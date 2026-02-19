# ============================================================
# RISK OVERVIEW PANEL - Violence Risk Analysis from Notes
# Styled to match Patient History Panel
# ============================================================

from __future__ import annotations

import re
from datetime import datetime, timedelta
from collections import defaultdict
from typing import Dict, List, Any
import io

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QScrollArea,
    QFrame, QSizeGrip, QSizePolicy, QToolTip, QMenu
)
from PySide6.QtCore import Qt, Signal, QPoint, QRect, QPropertyAnimation, QEasingCurve
from PySide6.QtGui import QColor, QPixmap, QCursor, QAction
from PySide6.QtWidgets import QMessageBox


def show_styled_message(parent, title: str, message: str, is_warning: bool = False):
    """Show a styled message box that works in dark mode."""
    msg = QMessageBox(parent)
    msg.setWindowTitle(title)
    msg.setText(message)
    msg.setIcon(QMessageBox.Warning if is_warning else QMessageBox.Information)
    msg.setStyleSheet("""
        QMessageBox {
            background-color: #2d2d2d;
        }
        QMessageBox QLabel {
            color: #ffffff;
            font-size: 13px;
        }
        QMessageBox QPushButton {
            background-color: #404040;
            color: #ffffff;
            border: 1px solid #555555;
            padding: 6px 20px;
            min-width: 70px;
            border-radius: 4px;
        }
        QMessageBox QPushButton:hover {
            background-color: #505050;
        }
    """)
    msg.exec()


# ============================================================
# RESIZABLE CHART CONTAINER - Fixed height with drag bar to resize
# ============================================================
class ResizableChartContainer(QWidget):
    """A container with horizontal scroll for charts and a drag bar to resize height."""

    def __init__(self, initial_height=150, min_height=80, max_height=500, parent=None):
        super().__init__(parent)
        self.min_height = min_height
        self.max_height = max_height
        self._dragging = False
        self._drag_start_y = 0
        self._orig_height = 0

        self.setFixedHeight(initial_height)
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Fixed)
        self.setMinimumWidth(1)

        # Main layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Scroll area for the chart (horizontal scroll only)
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal {
                background: rgba(0,0,0,0.1); height: 8px; border-radius: 4px;
            }
            QScrollBar::handle:horizontal {
                background: rgba(0,0,0,0.3); border-radius: 4px; min-width: 40px;
            }
        """)
        self.scroll.viewport().setStyleSheet("background: transparent;")
        layout.addWidget(self.scroll, 1)

        # Drag handle at bottom
        self.handle = QWidget()
        self.handle.setFixedHeight(12)
        self.handle.setCursor(Qt.SizeVerCursor)
        self.handle.setStyleSheet("""
            QWidget {
                background-color: rgba(100,100,100,0.3);
                border-radius: 3px;
                margin: 2px 40px;
            }
            QWidget:hover {
                background-color: rgba(100,150,200,0.5);
            }
        """)
        layout.addWidget(self.handle)

        # Install event filter on handle for drag
        self.handle.installEventFilter(self)

    def setWidget(self, widget):
        """Set the chart widget inside the scroll area."""
        self.scroll.setWidget(widget)

    def viewport(self):
        """Return the scroll area viewport for styling."""
        return self.scroll.viewport()

    def eventFilter(self, obj, event):
        """Handle drag events on the resize handle."""
        from PySide6.QtCore import QEvent

        if obj is self.handle:
            if event.type() == QEvent.MouseButtonPress:
                if event.button() == Qt.LeftButton:
                    self._dragging = True
                    self._drag_start_y = event.globalPosition().y()
                    self._orig_height = self.height()
                    return True

            elif event.type() == QEvent.MouseMove:
                if self._dragging:
                    dy = event.globalPosition().y() - self._drag_start_y
                    new_height = int(self._orig_height + dy)
                    new_height = max(self.min_height, min(self.max_height, new_height))
                    self.setFixedHeight(new_height)
                    return True

            elif event.type() == QEvent.MouseButtonRelease:
                if self._dragging:
                    self._dragging = False
                    return True

        return super().eventFilter(obj, event)

# Reuse shared components from patient history panel
from patient_history_panel_shared import CollapsibleSection, apply_macos_blur


# ============================================================
# RISK PATTERNS - Detailed subcategories from violence_risk_analyzer
# ============================================================

# Main categories with their subcategories, each containing severity levels
RISK_CATEGORIES = {
    "Verbal Aggression": {
        "color": "#9E9E9E",  # Light Grey
        "icon": "🗣️",
        "subcategories": {
            "Racial Abuse": {
                "severity": "high",
                "color": "#c0392b",
                "patterns": [
                    r'\b(black|white|asian|muslim|jew|n[i1]gg)\s*(cunt|bitch|bastard)\b',
                    r'\bracially\s+abusive\b',
                    r'\bracial\s+slur\b',
                ]
            },
            "Sexual/Homophobic Slurs": {
                "severity": "high",
                "color": "#8e44ad",
                "patterns": [
                    r'\b(puff|lesbian|queer|gay|fag)\b.{0,20}(insult|abus|call)',
                    r'(called?|calling)\s+.{0,15}(puff|lesbian|queer|fag)',
                ]
            },
            "Direct Insults": {
                "severity": "medium",
                "color": "#9b59b6",
                "patterns": [
                    r'(called?|calling)\s+(staff|them|him|her)?\s*.{0,10}(cunt|bitch|bastard|idiot|stupid)',
                    r'\bname[\s-]?calling\b',
                    r'\bcalling\s+staff\s+.{0,20}(cunt|names)\b',
                ]
            },
            "Swearing At": {
                "severity": "medium",
                "color": "#e67e22",
                "patterns": [
                    r'\b(fuck\s+off|told\s+.{0,10}fuck\s+off)',
                    r'(said|told|shouted)\s+.{0,15}fuck',
                    r'\bswearing\s+(at|towards)\b',
                ]
            },
            "Verbal Abuse": {
                "severity": "medium",
                "color": "#e74c3c",
                "patterns": [
                    r'\bverbally\s+(abusive|aggressive|hostile)\b',
                    r'\bverbal(ly)?\s+abus\w*',
                    r'\babusive\s+(language|towards|to\s+staff)\b',
                    r'\b(was|became|being)\s+abusive\b',
                ]
            },
            "Shouting": {
                "severity": "low",
                "color": "#f39c12",
                "patterns": [
                    r'\bshouting\s+(at|and)\b',
                    r'\bshouted\s+(at|and)\b',
                    r'\b(raised|raising)\s+(his|her)?\s*voice\b',
                    r'\bepisodes?\s+of\s+shouting\b',
                ]
            },
            "Threatening Language": {
                "severity": "high",
                "color": "#d35400",
                "patterns": [
                    r'\bthreateni?n?g?\s+(staff|peer|behaviour|language)\b',
                    r'\b(made?|making)\s+(a\s+)?threat\b',
                    r'(said|stated).{0,20}(kill|die|hurt|harm)',
                ]
            },
            "Spitting": {
                "severity": "high",
                "color": "#c0392b",
                "patterns": [
                    r'\bspat\s+(at|on|towards)\b',
                    r'\bspit(ting)?\s+(at|on|towards)\b',
                    r'\bspat\s+.{0,20}staff\b',
                ]
            },
            "Intimidation": {
                "severity": "high",
                "color": "#8e44ad",
                "patterns": [
                    r'\bintimidati\w+',
                    r'\b(charged?|charging)\s+(at|towards)\b',
                    r'\b(squared?|squaring)\s+up\b',
                    r'\b(got|getting)\s+in\s+.{0,10}face\b',
                ]
            },
        },
        "patterns": []  # Will be populated from subcategories
    },
    "Physical Aggression": {
        "color": "#b71c1c",
        "icon": "👊",
        "subcategories": {
            "Assault on Staff": {
                "severity": "high",
                "color": "#b71c1c",
                "patterns": [
                    r'\b(punch\w*|kick\w*|hit|slap\w*|struck|attack\w*)\s+(a\s+)?(staff|nurse|hca|doctor|member\s+of\s+staff)\b',
                    r'\bassault\w*\s+(a\s+)?(staff|nurse|hca|doctor)\b',
                    r'\b(head.?butt\w*|headbutt\w*)\s+(a\s+)?(staff|nurse)\b',
                    r'\b(bit|biting|bitten)\s+(a\s+)?(staff|nurse)\b',
                    r'\b(scratch\w*|claw\w*)\s+(at\s+)?(staff|nurse)\b',
                    r'\b(push\w*|shov\w*)\s+(a\s+)?(staff|nurse)\b',
                ]
            },
            "Assault on Peer": {
                "severity": "high",
                "color": "#c62828",
                "patterns": [
                    r'\b(punch\w*|kick\w*|hit|slap\w*|struck|attack\w*)\s+(a\s+)?(peer|patient|another\s+patient)\b',
                    r'\bassault\w*\s+(a\s+)?(peer|patient|another)\b',
                    r'\bphysical\s+altercation\b',
                    r'\b(head.?butt\w*|headbutt\w*)\s+(a\s+)?(peer|patient)\b',
                    r'\b(bit|biting|bitten)\s+(a\s+)?(peer|patient)\b',
                ]
            },
            "Physical Aggression": {
                "severity": "high",
                "color": "#d32f2f",
                "patterns": [
                    r'\bphysical(ly)?\s+aggress\w*',
                    r'\blashed\s+out\s+(at|physically)\b',
                    r'\b(became|becoming|was)\s+physically\s+violent\b',
                    r'\bviolent\s+(outburst|episode|incident)\b',
                ]
            },
            "Restraint Required": {
                "severity": "high",
                "color": "#ef5350",
                "patterns": [
                    r'\b(restrain\w*|restraint)\s+(was\s+)?(required|needed|used|applied)\b',
                    r'\b(prone|supine)\s+restraint\b',
                    r'\brequired\s+(physical\s+)?intervention\b',
                    r'\b(rapi?d\s+tranquil|rt\s+administered|given\s+rt|rt\s+given)',
                    r'\bprn\s+(medication\s+)?(given|administered)\s+(for|due\s+to)\s+(aggression|violence)',
                ]
            },
            "Attempted Violence": {
                "severity": "medium",
                "color": "#ef5350",
                "patterns": [
                    r'\battempt\w*\s+to\s+(punch|kick|hit|slap|attack|assault|strike)\b',
                    r'\b(tried|trying)\s+to\s+(punch|kick|hit|slap|attack|assault)\b',
                    r'\b(lunge\w*|swing\w*|swung)\s+(at|towards)\s+(staff|peer|patient)\b',
                    r'\b(raised|clenched)\s+(his|her)\s+fist\b',
                ]
            },
            "Use of Weapon/Object": {
                "severity": "high",
                "color": "#c62828",
                "patterns": [
                    r'\b(used|using|brandish\w*)\s+(a\s+)?(weapon|knife|blade|sharp|object)\b',
                    r'\b(threw|thrown|throwing)\s+(a\s+)?(heavy|hard)\s+object\b',
                    r'\b(arm\w*|armed)\s+(himself|herself)\s+with\b',
                ]
            },
        },
        "patterns": []
    },
    "Property Damage": {
        "color": "#e53935",
        "icon": "🔨",
        "subcategories": {
            "Breaking Items": {
                "severity": "medium",
                "color": "#e53935",
                "patterns": [
                    r'\b(broke|broken|breaking|smash\w*|damag\w*)\s+(the\s+)?(window|door|furniture|tv|television|chair|table)\b',
                    r'\bdamag\w+\s+(to\s+)?(property|furniture|equipment)\b',
                ]
            },
            "Punching/Kicking Objects": {
                "severity": "medium",
                "color": "#ef5350",
                "patterns": [
                    r'\b(punch\w*|kick\w*|hit)\s+(the\s+)?(wall|door|window)\b',
                ]
            },
            "Throwing Objects": {
                "severity": "medium",
                "color": "#f44336",
                "patterns": [
                    r'\b(threw|thrown|throwing)\s+.{0,15}(furniture|chair|table|object|item)\b',
                    r'\b(overturn\w*|upend\w*|flip\w*)\s+(the\s+)?(table|chair|furniture|bed)\b',
                ]
            },
            "Room Destruction": {
                "severity": "high",
                "color": "#c62828",
                "patterns": [
                    r'\b(destroy\w*|wreck\w*)\s+(his|her|the)\s+(room|property|belongings)\b',
                    r'\b(vandal\w*|trash\w*)\s+(his|her|the)\s+(room|property)\b',
                ]
            },
        },
        "patterns": []
    },
    "Self-Harm": {
        "color": "#ff5722",
        "icon": "⚠️",
        "subcategories": {
            "Cutting": {
                "severity": "high",
                "color": "#ff5722",
                "patterns": [
                    r'\b(he|she|patient|resident)\s+(cut|cuts|cutting)\s+(his|her)\s+(arm|wrist|leg|body|skin|face)\b',
                    r'\bself[\s-]?cut\w*\s+(today|this|during)\b',
                    r'\b(slash|slashing|slashed)\s+(his|her)\s+(wrist|arm)\b',
                ]
            },
            "Head Banging": {
                "severity": "high",
                "color": "#e65100",
                "patterns": [
                    r'\b(he|she|patient)\s+(was|started|began)?\s*bang\w*\s+(his|her)\s*head\b',
                    r'\bhead[\s-]?bang\w*\s+(incident|episode|observed|witnessed)\b',
                    r'\b(he|she)\s+hit\s+(his|her)\s+head\s+(against|on|into)\s+(wall|door|floor)\b',
                ]
            },
            "Hitting Self": {
                "severity": "high",
                "color": "#f57c00",
                "patterns": [
                    r'\b(he|she|patient)\s+(hit|hitting|hits|struck|punched|punching)\s+(himself|herself)\b',
                    r'\b(slap|slapping|slapped)\s+(himself|herself|his|her)\s+(face|head)\b',
                ]
            },
            "Ligature": {
                "severity": "high",
                "color": "#d84315",
                "patterns": [
                    r'\b(made|tied|created|found\s+with|attempted)\s+(a\s+)?ligature\b',
                    r'\bligature\s+(incident|attempt|found|discovered)\b',
                    r'\b(attempted?|attempt\w*|tried)\s+to\s+hang\s+(himself|herself)\b',
                ]
            },
            "Overdose": {
                "severity": "high",
                "color": "#bf360c",
                "patterns": [
                    r'\b(he|she|patient)\s+(took|taken|has\s+taken)\s+(an?\s+)?overdose\b',
                    r'\boverdose\s+(incident|attempt|today|this)\b',
                    r'\b(swallowed|ingested)\s+(multiple|extra|excess)\s+(tablet|pill|medication)\b',
                ]
            },
            "Self-Harm Threat": {
                "severity": "medium",
                "color": "#ff8a65",
                "patterns": [
                    r'\b(he|she|patient)\s+(threatened?|threatening)\s+to\s+(harm|hurt|cut|kill)\s+(himself|herself)\b',
                    r'\b(threatened?|threatening)\s+to\s+self[\s-]?harm\b',
                ]
            },
            "Self-Harm Act": {
                "severity": "high",
                "color": "#dd2c00",
                "patterns": [
                    r'\b(he|she|patient)\s+self[\s-]?harmed\b',
                    r'\bself[\s-]?harm\s+(incident|episode|act)\s+(today|this|occurred|reported)\b',
                    r'\b(engaged?|engaging)\s+in\s+self[\s-]?harm\w*\b',
                ]
            },
        },
        "patterns": []
    },
    "Sexual Behaviour": {
        "color": "#00BCD4",
        "icon": "🚫",
        "subcategories": {
            "Sexual Comments": {
                "severity": "medium",
                "color": "#e91e63",
                "patterns": [
                    r'\b(sexual|inappropriate)\s+(comment|remark)\b',
                    r'\bcomment\w*\s+(of\s+)?(a\s+)?sexual\s+nature\b',
                    r'\b(lewd|obscene|vulgar)\s+(comment|remark|language)\b',
                ]
            },
            "Sexual Touching": {
                "severity": "high",
                "color": "#c2185b",
                "patterns": [
                    r'\b(inappropriately?\s+)?touch\w*\s+(staff|breast|buttock|bottom|bum|chest|groin|thigh)\b',
                    r'\b(tried?|attempt\w*|trying)\s+(to\s+)?(touch|grab|grope)\b',
                    r'\b(grope|groping|groped)\b',
                    r'\b(unwanted|inappropriate|unsolicited)\s+touch\b',
                ]
            },
            "Exposure": {
                "severity": "high",
                "color": "#880e4f",
                "patterns": [
                    r'\bexpos(ed?|ing)\s+(himself|herself|genitals?|private)\b',
                    r'\bexpos\w+\s+inappropriately\b',
                    r'\b(flash\w*|flashing)\s+(himself|herself|staff|peer)\b',
                ]
            },
            "Public Masturbation": {
                "severity": "high",
                "color": "#ad1457",
                "patterns": [
                    r'\bmasturbat\w+',
                    r'\b(observed?|caught|found)\s+.{0,20}masturbat',
                ]
            },
            "Walking Naked": {
                "severity": "high",
                "color": "#d81b60",
                "patterns": [
                    r'\b(walk\w*|came|went)\s+.{0,10}naked\b',
                    r'\bnaked\s+(in|on|around)\s+(the\s+)?(floor|corridor|lounge|ward|public)\b',
                ]
            },
            "Sexual Disinhibition": {
                "severity": "medium",
                "color": "#f06292",
                "patterns": [
                    r'\bsexual(ly)?\s+disinhibit\w*',
                    r'\b(overly?\s+)?flirtatious\b',
                    r'\b(sexual|inappropriate)\s+gesture\b',
                ]
            },
            "Sexual Advances": {
                "severity": "high",
                "color": "#ec407a",
                "patterns": [
                    r'\b(sexual|inappropriate)\s+advance\b',
                    r'\bproposition\w*',
                    r'\b(tried?|attempt\w*)\s+(to\s+)?kiss\b',
                ]
            },
        },
        "patterns": []
    },
    "Bullying/Exploitation": {
        "color": "#795548",
        "icon": "😠",
        "subcategories": {
            "Bullying Peer": {
                "severity": "high",
                "color": "#795548",
                "patterns": [
                    r'\b(was|been|observed|witnessed|seen)\s+(to\s+)?(be\s+)?(bully|bullying)\b',
                    r'\b(bully|bullying)\s+(peer|patient|another|other)\b',
                    r'\b(pick|picking|picked)\s+on\s+(peer|patient|another)\b',
                ]
            },
            "Targeting Vulnerable": {
                "severity": "high",
                "color": "#5d4037",
                "patterns": [
                    r'\b(target|targeting|targeted)\s+(vulnerable|other)?\s*(patient|peer)\b',
                ]
            },
            "Taking Items": {
                "severity": "medium",
                "color": "#8d6e63",
                "patterns": [
                    r'\b(took|taking|taken|steal|stealing|stolen)\s+.{0,15}(food|cigarette|vape|belonging|item|money)\s+(from|of)\s+(peer|patient)\b',
                    r'\b(demand|demanding|demanded)\s+.{0,15}(food|cigarette|money|item)\s+from\b',
                ]
            },
            "Financial Exploitation": {
                "severity": "high",
                "color": "#6d4c41",
                "patterns": [
                    r'\b(extort|extorting|extorted)\b',
                    r'\b(exploit|exploiting)\s+(other|peer|patient)\b',
                    r'\b(borrow|borrowing|borrowed)\s+money\s+from\s+(peer|patient)\b',
                ]
            },
            "Intimidation of Peers": {
                "severity": "high",
                "color": "#4e342e",
                "patterns": [
                    r'\b(intimidat\w+)\s+(peer|patient|another|other)\b',
                    r'\bphysically\s+threatening\s+to\s+other\s+patient\b',
                ]
            },
            "Coercion": {
                "severity": "high",
                "color": "#3e2723",
                "patterns": [
                    r'\b(coerce|coercing|coerced)\s+(peer|patient|another)\b',
                    r'\b(force|forcing|forced)\s+(peer|patient|another)\s+to\b',
                    r'\b(pressure|pressuring|pressured)\s+(peer|patient|another)\b',
                ]
            },
        },
        "patterns": []
    },
    "Self-Neglect": {
        "color": "#607d8b",
        "icon": "🧹",
        "subcategories": {
            "Unkempt Appearance": {
                "severity": "medium",
                "color": "#607d8b",
                "patterns": [
                    r'\b(looked?|appear\w*|present\w*)\s+(very\s+)?(unkempt|dishevelled|unwashed|neglected)\b',
                    r'\b(he|she)\s+(was|looked?|appeared?)\s+(unkempt|dishevelled)\b',
                ]
            },
            "Dirty Clothes": {
                "severity": "medium",
                "color": "#78909c",
                "patterns": [
                    r'\bclothes?\s+(were?|was|is|are)\s+(dirty|soiled|stained)\b',
                    r'\bwearing\s+(same|dirty|soiled)\s+(clothes?|clothing)\b',
                ]
            },
            "Body Odour": {
                "severity": "medium",
                "color": "#546e7a",
                "patterns": [
                    r'\bbody\s+odou?r\b',
                    r'\b(malodorous|strong\s+smell)\b',
                    r'\bsmell\w*\s+(of\s+)?(urine|faec|body)\b',
                ]
            },
            "Refused Self-Care": {
                "severity": "medium",
                "color": "#455a64",
                "patterns": [
                    r'\b(declined?|refused?|reluctant)\s+to\s+(shower|wash|bathe|change)\b',
                    r'\b(declined?|refused?)\s+(shower|personal\s+care|self.?care)\b',
                ]
            },
            "Requires Prompting": {
                "severity": "low",
                "color": "#37474f",
                "patterns": [
                    r'\b(prompting|prompted?)\s+(to|for)\s+(shower|wash|self.?care|personal\s+care)\b',
                    r'\b(require\w*|need\w*)\s+prompt\w*\s+(to|for)\s+(shower|wash|personal|hygiene)\b',
                ]
            },
            "Poor Room State": {
                "severity": "low",
                "color": "#263238",
                "patterns": [
                    r'\broom\s+management\s+(was\s+)?poor\b',
                    r'\b(room|bedroom)\s+(was|is|in)\s+(a\s+)?(dirty|mess|poor\s+state)\b',
                ]
            },
            "Poor Dietary Intake": {
                "severity": "medium",
                "color": "#90a4ae",
                "patterns": [
                    r'\bpoor\s+(dietary|fluid|food)\s+(and\s+fluid\s+)?intake\b',
                    r'\b(declined?|refused?)\s+(all\s+)?(food|meals?|to\s+eat)\b',
                ]
            },
        },
        "patterns": []
    },
    "AWOL/Absconding": {
        "color": "#f57c00",
        "icon": "🚪",
        "subcategories": {
            "AWOL": {
                "severity": "high",
                "color": "#f57c00",
                "patterns": [
                    r'\bawol\b',
                    r'\babsent\s+without\s+leave\b',
                    r'\babscond\w*',
                ]
            },
            "Failed to Return": {
                "severity": "medium",
                "color": "#ff9800",
                "patterns": [
                    r'\bfailed\s+to\s+return\b',
                    r'\bdid\s+not\s+return\s+from\s+leave\b',
                ]
            },
            "Escape Attempt": {
                "severity": "high",
                "color": "#e65100",
                "patterns": [
                    r'\b(escaped?|escaping)\s+from\b',
                    r'\bleft\s+without\s+permission\b',
                    r'\bbreach\s+of\s+leave\b',
                ]
            },
        },
        "patterns": []
    },
    "Substance Misuse": {
        "color": "#9c27b0",
        "icon": "💊",
        "subcategories": {
            "Positive Drug Test": {
                "severity": "high",
                "color": "#6a1b9a",
                "patterns": [
                    r'\b(tested|test)\s+(came\s+back\s+)?positive\s+(for\s+)?(drug|cannabis|thc|cocaine|amphet|opiates?|benzo)\b',
                    r'\bpositive\s+(drug|uds|urine)\s+(test|screen|result)\b',
                    r'\buds\s+(was\s+|came\s+back\s+)?positive\b',
                ]
            },
            "Smelling of Substances": {
                "severity": "medium",
                "color": "#7b1fa2",
                "patterns": [
                    r'\b(smell\w*|smelt)\s+(of|like)\s+(cannabis|alcohol|weed|spice|marijuana)\b',
                    r'\b(alcohol|intoxicat)\w*\s+(on\s+breath|smelt|smell|detected)\b',
                ]
            },
            "Appeared Intoxicated": {
                "severity": "high",
                "color": "#8e24aa",
                "patterns": [
                    r'\b(appear\w*|seem\w*|present\w*)\s+(to\s+be\s+)?(intoxicated|drunk|under\s+the\s+influence)\b',
                    r'\b(was|were|is)\s+(visibly\s+)?(intoxicated|drunk)\b',
                ]
            },
            "Admitted Substance Use": {
                "severity": "high",
                "color": "#9c27b0",
                "patterns": [
                    r'\b(admitted|disclosed|confessed)\s+(to\s+)?(using|smoking|taking|drinking)\s+(cannabis|spice|drugs|alcohol|cocaine)\b',
                ]
            },
            "Found with Substances": {
                "severity": "high",
                "color": "#ab47bc",
                "patterns": [
                    r'\bfound\s+(with\s+)?(drugs|substances|cannabis|alcohol)\b',
                    r'\b(suspected|illicit)\s+substance\b',
                ]
            },
        },
        "patterns": []
    },
    "Non-Compliance": {
        "color": "#dc2626",
        "icon": "✋",
        "subcategories": {
            "Refused Medication": {
                "severity": "medium",
                "color": "#dc2626",
                "patterns": [
                    r'\b(declined?|refused?)\s+(his|her|all)?\s*meds?\b',
                    r'\b(declined?|refused?)\s+(depot|clozapine|treatment)\b',
                    r'\bdid\s+not\s+(take|have)\s+meds?\b',
                ]
            },
            "Non-Compliance": {
                "severity": "medium",
                "color": "#ef4444",
                "patterns": [
                    r'\bnon[\s-]?compli\w*',
                ]
            },
            "Refused to Engage": {
                "severity": "low",
                "color": "#f87171",
                "patterns": [
                    r'\b(declined?|refused?)\s+to\s+(engage|attend|participate)\b',
                ]
            },
        },
        "patterns": []
    },
}

# Populate main patterns list from subcategories for backwards compatibility
for cat_name, cat_data in RISK_CATEGORIES.items():
    if "subcategories" in cat_data:
        all_patterns = []
        for subcat_data in cat_data["subcategories"].values():
            all_patterns.extend(subcat_data["patterns"])
        cat_data["patterns"] = all_patterns

# False positive exclusion patterns (checked BEFORE match)
FALSE_POSITIVE_PATTERNS = [
    r'(was|were|has|had|have)\s+not\s+(been\s+)?(physically\s+)?',
    r"(wasn't|weren't|hasn't|hadn't|haven't)\s+(been\s+)?",
    r'not\s+(physically\s+)?(aggressive|violent|agitated|abusive|threatening)',
    r'no\s+(physical\s+)?(aggression|violence|agitation|abuse|threats?)',
    r'(without|nil|no)\s+(any\s+)?(aggression|violence|incident)',
    r'no\s+(episodes?\s+of\s+)?',
    r'(there\s+was|there\s+were)\s+no\s+',
    r'(risk of|at risk of|risk assessment|level of risk)',
    r'(potential|possibility|likelihood)\s+(of|for)\s+',
    r'remains\s+(a\s+)?(risk|unpredictable|concern)',
    r'(background|history of|past history|previous)',
    r'(remained|remains|stayed|was)\s+(calm|settled|pleasant|cooperative)',
    # Conditional/hypothetical statements
    r'\bif\s+(there\s+is\s+)?(any\s+)?(self[\s-]?harm|violent|violence|aggression|aggressive|incident|harm)',
    r'\bif\s+(he|she|patient|they)\s+(is|are|becomes?|were|was)\s+',
    r'\bin\s+case\s+of\s+',
    r'\bin\s+the\s+event\s+of\s+',
    r'\bshould\s+(there\s+be|he|she|they)\s+',
    r'\bto\s+be\s+called\s+if\s+',
    r'\bcontact\s+.{0,20}\s+if\s+',
    r'\bplan\s+(is\s+)?(to|for)\s+',
]


def _has_negative_context(text: str, match_start: int, match_end: int) -> bool:
    """Check if a match has negative context (nil, no, denied, etc.)."""
    # Get context around the match (50 chars before and after)
    context_start = max(0, match_start - 50)
    context_end = min(len(text), match_end + 50)

    before_text = text[context_start:match_start].lower()
    after_text = text[match_end:context_end].lower()

    # Check for negative patterns before the match
    negative_before = [
        r'\b(no|nil|none|denies?|denied|without|lacks?)\s*$',
        r'\b(no|nil|none|denies?|denied|without|lacks?)\s+(any|all|the|a)?\s*$',
        r'\b(no\s+evidence|no\s+history|no\s+signs?|no\s+indication)\s+of\s*$',
        r'\b(has\s+not|did\s+not|does\s+not|hasn\'t|didn\'t|doesn\'t)\s*$',
        r'\b(not\s+noted|not\s+reported|not\s+observed)\s*$',
        r'\bdenied\s+(any|all)?\s*(thoughts?\s+of|ideation\s+of|intent\s+to|intention\s+to)?\s*$',
        # Additional patterns for common clinical negations
        r'\b(not|never)\s+express(ing|ed)?\s*(any)?\s*$',
        r'\bno\s+(current|recent|active|new)?\s*(thoughts?|episodes?|ideation|urges?|intent)\s+(of|to)\s*$',
        r'\b(doesn\'t|does\s+not|don\'t|do\s+not)\s+have\s+(\w+\s+)*$',
        r'\b(there\s+)?(were|was|are|is)\s+no\s+(episode|evidence|history|indication)s?\s+(of)?\s*$',
        r'\bno\s+\w+\s+of\s*$',  # "no episode of", "no thoughts of", etc.
        r'\bwith\s+no\s*$',  # "with no self-harming"
        # Handle "or" clauses - negation earlier in phrase carries through
        r'\bor\s+(urges?\s+to|thoughts?\s+of|intent\s+to)?\s*$',
        # "was not X", "were not X" patterns
        r'\b(was|were|is|are)\s+not\s*$',
    ]

    for pattern in negative_before:
        if re.search(pattern, before_text):
            return True

    # Check if negation appears earlier in the broader context (handles "X or Y" constructs)
    broader_negation = [
        r'\b(no|not|nil|none|denied|denies|without)\b.{0,40}\bor\b',
        r'\b(doesn\'t|does\s+not|don\'t)\s+have\b',
    ]
    for pattern in broader_negation:
        if re.search(pattern, before_text):
            return True

    # Check for negative patterns after the match
    negative_after = [
        r'^\s*(nil|none|denied|not\s+noted|not\s+reported)\b',
        r'^\s*-\s*(nil|no|none|denied)\b',
        r'^\s*(were|was|are|is)?\s*(not\s+present|not\s+identified|not\s+expressed|absent)\b',
        r'^\s*(were|was|are|is)?\s*not\s+(noted|reported|observed|evident|identified)\b',
        r'^\s*\w*\s*(were|was|are|is)?\s*(not\s+present|absent)\b',
        r'^\s*(thoughts?|ideation)?\s*(were|was|are|is)?\s*(not\s+present|absent|not\s+expressed|denied)\b',
        # "X was not required/needed" patterns
        r'^\s*(was|were|is|are)?\s*not\s+(required|needed|necessary|indicated)\b',
    ]

    for pattern in negative_after:
        if re.search(pattern, after_text):
            return True

    # Check full context for assessment/documentation language
    full_context = text[context_start:context_end].lower()
    assessment_patterns = [
        r'\brisk\s+(assessment|screen|factor)',
        r'\b(asked|enquired|assessed)\s+about\b',
        r'\bqueried\s+(re|regarding|about)\b',
    ]

    for pattern in assessment_patterns:
        if re.search(pattern, full_context):
            # Only exclude if also has negative indicator
            if re.search(r'\b(no|nil|denied|denies|negative)\b', full_context):
                return True

    return False


def _normalise_date(d):
    """Convert any date format to datetime."""
    if d is None:
        return None
    if isinstance(d, datetime):
        return d
    try:
        import pandas as pd
        dt = pd.to_datetime(d, errors="coerce", dayfirst=True)
        if pd.isna(dt):
            return None
        return dt.to_pydatetime()
    except Exception:
        return None


def is_false_positive(text: str, match_start: int) -> bool:
    """Check if a match is a false positive based on context."""
    context_start = max(0, match_start - 60)
    context = text[context_start:match_start].lower()

    for fp_pattern in FALSE_POSITIVE_PATTERNS:
        if re.search(fp_pattern, context, re.IGNORECASE):
            return True
    return False


def highlight_matches(text: str, patterns: List[str]) -> str:
    """Highlight matched patterns in text with yellow background."""
    highlighted = text
    matches_found = []
    text_lower = text.lower()

    for pattern in patterns:
        try:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                if not is_false_positive(text, match.start()):
                    # Also check for negative context
                    if not _has_negative_context(text_lower, match.start(), match.end()):
                        matches_found.append((match.start(), match.end(), match.group()))
        except:
            pass

    # Sort by position descending to replace from end first
    matches_found.sort(key=lambda x: x[0], reverse=True)

    for start, end, matched_text in matches_found:
        highlighted = (
            highlighted[:start] +
            f"<span style='background-color: #FFEB3B; color: #000; padding: 2px 4px; border-radius: 3px;'>{matched_text}</span>" +
            highlighted[end:]
        )

    return highlighted


def analyze_notes_for_risk(notes: List[Dict]) -> Dict[str, Any]:
    """Analyze notes for risk incidents and return summary with subcategories and severity."""
    results = {
        "total_notes": len(notes),
        "notes_with_incidents": 0,
        "categories": {},
        "timeline": [],
        "monthly_counts": defaultdict(lambda: defaultdict(int)),
        "severity_counts": {"high": 0, "medium": 0, "low": 0},
    }

    # Initialize categories with subcategories
    for cat_name, cat_config in RISK_CATEGORIES.items():
        results["categories"][cat_name] = {
            "color": cat_config["color"],
            "icon": cat_config["icon"],
            "patterns": cat_config["patterns"],
            "count": 0,
            "incidents": [],
            "subcategories": {},
        }
        # Initialize subcategories if present
        if "subcategories" in cat_config:
            for subcat_name, subcat_config in cat_config["subcategories"].items():
                results["categories"][cat_name]["subcategories"][subcat_name] = {
                    "color": subcat_config["color"],
                    "severity": subcat_config["severity"],
                    "count": 0,
                    "incidents": [],
                }

    for note in notes:
        text = note.get("text", "") or note.get("content", "") or note.get("body", "")
        date = _normalise_date(note.get("date") or note.get("datetime"))

        if not text:
            continue

        text_lower = text.lower()
        note_had_incident = False
        matched_in_note = set()  # Track what we've matched to avoid duplicates

        for cat_name, cat_config in RISK_CATEGORIES.items():
            # Check subcategories first for more specific matching
            if "subcategories" in cat_config:
                for subcat_name, subcat_config in cat_config["subcategories"].items():
                    for pattern in subcat_config["patterns"]:
                        try:
                            match = re.search(pattern, text_lower)
                            if match:
                                # Skip if same match already found in this note for this category
                                match_key = (cat_name, match.group())
                                if match_key in matched_in_note:
                                    continue

                                if is_false_positive(text, match.start()):
                                    continue

                                # Check for negative context (nil, no, denied, etc.)
                                if _has_negative_context(text_lower, match.start(), match.end()):
                                    continue

                                matched_in_note.add(match_key)
                                severity = subcat_config["severity"]

                                # Update main category
                                results["categories"][cat_name]["count"] += 1
                                results["categories"][cat_name]["incidents"].append({
                                    "date": date,
                                    "full_text": text,
                                    "matched": match.group(),
                                    "subcategory": subcat_name,
                                    "severity": severity,
                                })

                                # Update subcategory
                                results["categories"][cat_name]["subcategories"][subcat_name]["count"] += 1
                                results["categories"][cat_name]["subcategories"][subcat_name]["incidents"].append({
                                    "date": date,
                                    "full_text": text,
                                    "matched": match.group(),
                                    "severity": severity,
                                })

                                # Update severity counts
                                results["severity_counts"][severity] += 1

                                results["timeline"].append((date, cat_name, text, subcat_name, severity))

                                if date:
                                    month_key = date.strftime("%Y-%m")
                                    results["monthly_counts"][month_key][cat_name] += 1

                                note_had_incident = True
                                break  # Move to next subcategory after finding a match
                        except Exception:
                            pass

        if note_had_incident:
            results["notes_with_incidents"] += 1

    results["timeline"].sort(key=lambda x: x[0] or datetime.min, reverse=True)

    return results


def create_risk_timeline_visual(results: Dict) -> tuple:
    """Create a visual timeline chart showing risk levels over time.

    Returns:
        tuple: (image_bytes, timeline_info) with bar positions for interactivity
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        from io import BytesIO

        if not results.get("monthly_counts"):
            return None, {}

        monthly_counts = results["monthly_counts"]
        categories = results.get("categories", {})
        months = sorted(monthly_counts.keys())
        if len(months) < 2:
            return None, {}

        # Calculate risk level for each month
        month_data = []
        for month in months:
            counts = monthly_counts[month]
            total = sum(counts.values())

            # Categorize activity level
            if total == 0:
                level = 0  # quiet
                color = '#2d5a3d'  # dark green
                level_name = "Quiet"
            elif total <= 3:
                level = 1  # low
                color = '#22c55e'  # green
                level_name = "Low"
            elif total <= 8:
                level = 2  # moderate
                color = '#f59e0b'  # amber
                level_name = "Moderate"
            elif total <= 15:
                level = 3  # elevated
                color = '#f97316'  # orange
                level_name = "Elevated"
            else:
                level = 4  # high
                color = '#ef4444'  # red
                level_name = "High"

            # Get top categories for this month
            top_cats = sorted(counts.items(), key=lambda x: x[1], reverse=True)[:3]

            month_data.append({
                "month": month,
                "level": level,
                "level_name": level_name,
                "color": color,
                "total": total,
                "counts": dict(counts),
                "top_cats": top_cats
            })

        # Create figure - horizontal timeline
        fig_width = max(12, len(months) * 0.4)
        fig, ax = plt.subplots(figsize=(fig_width, 2.5))
        fig.patch.set_facecolor('#1a1a2e')
        ax.set_facecolor('#1a1a2e')

        # Draw risk level bars and track positions
        bar_height = 0.6
        dpi = 100
        bar_positions = []

        for i, data in enumerate(month_data):
            ax.barh(0, 1, left=i, height=bar_height, color=data["color"], edgecolor='#333', linewidth=0.5)

            # Add incident count on bar if > 0
            if data["total"] > 0:
                ax.text(i + 0.5, 0, str(data["total"]), ha='center', va='center',
                       fontsize=8, color='white', fontweight='bold')

        # X-axis labels (months)
        ax.set_xlim(0, len(months))
        ax.set_ylim(-0.5, 0.8)

        # Show month labels (every few months for readability)
        step = max(1, len(months) // 12)
        tick_positions = list(range(0, len(months), step))
        tick_labels = []
        for i in tick_positions:
            try:
                dt = datetime.strptime(months[i], "%Y-%m")
                tick_labels.append(dt.strftime("%b '%y"))
            except:
                tick_labels.append(months[i])

        ax.set_xticks([p + 0.5 for p in tick_positions])
        ax.set_xticklabels(tick_labels, fontsize=9, color='#AAA', rotation=45, ha='right')

        # Remove y-axis
        ax.set_yticks([])
        ax.set_ylabel('')

        # Title
        ax.set_title('Risk Level Timeline', fontsize=11, color='white', pad=10, loc='left')

        # Legend
        legend_elements = [
            mpatches.Patch(facecolor='#2d5a3d', edgecolor='#555', label='Quiet (0)'),
            mpatches.Patch(facecolor='#22c55e', edgecolor='#555', label='Low (1-3)'),
            mpatches.Patch(facecolor='#f59e0b', edgecolor='#555', label='Moderate (4-8)'),
            mpatches.Patch(facecolor='#f97316', edgecolor='#555', label='Elevated (9-15)'),
            mpatches.Patch(facecolor='#ef4444', edgecolor='#555', label='High (16+)'),
        ]
        ax.legend(handles=legend_elements, loc='upper right', fontsize=7,
                 facecolor='#2a2a3e', edgecolor='#555', labelcolor='white',
                 ncol=5, framealpha=0.9)

        # Remove spines
        for spine in ax.spines.values():
            spine.set_visible(False)

        plt.tight_layout()

        # Save to bytes first
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, facecolor='#1a1a2e',
                   edgecolor='none', bbox_inches='tight', pad_inches=0.1)

        # Get the actual saved image dimensions
        buf.seek(0)
        from PIL import Image
        img = Image.open(buf)
        img_width, img_height = img.size
        buf.seek(0)

        # Calculate bar positions based on image dimensions
        # The chart area is roughly the middle portion of the image
        # Account for title, legend, and axis labels
        left_margin = 50  # Approximate left margin for y-axis area
        right_margin = 30
        top_margin = 45   # Title + legend area
        bottom_margin = 50  # X-axis labels

        chart_width = img_width - left_margin - right_margin
        chart_height = img_height - top_margin - bottom_margin

        bar_width = chart_width / len(months)
        bar_pixel_height = chart_height * 0.5  # Bar is 0.6 of 1.3 total y range

        for i, data in enumerate(month_data):
            x_left = left_margin + i * bar_width
            x_right = left_margin + (i + 1) * bar_width
            y_top = top_margin + chart_height * 0.2  # Top of bar area
            y_bottom = top_margin + chart_height * 0.7  # Bottom of bar area

            bar_positions.append({
                'index': i,
                'month': data['month'],
                'x_left': x_left,
                'x_right': x_right,
                'y_top': y_top,
                'y_bottom': y_bottom,
                'total': data['total'],
                'level_name': data['level_name'],
                'counts': data['counts'],
                'top_cats': data['top_cats']
            })
        plt.close(fig)
        buf.seek(0)

        timeline_info = {
            'bar_positions': bar_positions,
            'months': months,
            'month_data': month_data
        }

        return buf.read(), timeline_info

    except Exception as e:
        print(f"Error creating risk timeline: {e}")
        import traceback
        traceback.print_exc()
        return None, {}


def create_timeline_chart(monthly_counts: Dict, categories_data: Dict,
                          y_max: int = None, color_override: str = None) -> tuple:
    """Create a clustered column chart showing incidents over time by month.

    Args:
        monthly_counts: Dict of {month_key: {category: count}}
        categories_data: Dict of category info from analyze_notes_for_risk
        y_max: Fixed Y-axis maximum (keeps scale consistent across filtered views)
        color_override: Hex color to use for all bars (e.g. the filtered category's color)

    Returns:
        tuple: (image_bytes, timeline_info) where timeline_info contains metadata
               for interactive tooltips
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np

        if not monthly_counts:
            return None, {}

        # Get sorted months
        months = sorted(monthly_counts.keys())
        if len(months) < 2:
            return None, {}

        # Get categories with data
        active_cats = [name for name, info in categories_data.items() if info["count"] > 0]
        if not active_cats:
            return None, {}

        # Limit to top 6 categories for readability
        cat_totals = [(cat, sum(monthly_counts[m].get(cat, 0) for m in months)) for cat in active_cats]
        cat_totals.sort(key=lambda x: x[1], reverse=True)
        top_cats = [c[0] for c in cat_totals[:6]]

        # Distinct colors for chart (very different from each other)
        CHART_COLORS = [
            '#e74c3c',  # Red
            '#3498db',  # Blue
            '#f39c12',  # Orange
            '#9b59b6',  # Purple
            '#1abc9c',  # Teal
            '#27ae60',  # Green
            '#e91e63',  # Pink
            '#00bcd4',  # Cyan
        ]

        # Build data matrix
        data_matrix = []
        colors = []
        for i, cat in enumerate(top_cats):
            row = [monthly_counts[m].get(cat, 0) for m in months]
            data_matrix.append(row)
            if color_override:
                colors.append(color_override)
            else:
                # Use the category's own color if available, else fall back to chart palette
                cat_color = categories_data.get(cat, {}).get("color")
                colors.append(cat_color if cat_color else CHART_COLORS[i % len(CHART_COLORS)])

        # Create figure
        fig_width, fig_height = 12, 4
        dpi = 100
        fig, ax = plt.subplots(figsize=(fig_width, fig_height))
        fig.patch.set_facecolor('#1a1a2e')
        ax.set_facecolor('#1a1a2e')

        x = np.arange(len(months))
        width = 0.8 / len(top_cats)

        all_bars = []
        for i, (cat, row, color) in enumerate(zip(top_cats, data_matrix, colors)):
            offset = (i - len(top_cats)/2 + 0.5) * width
            bars = ax.bar(x + offset, row, width, label=cat, color=color)
            all_bars.append((cat, bars))

        # Calculate time span in years
        first_year = int(months[0][:4])
        last_year = int(months[-1][:4])
        year_span = last_year - first_year + 1

        # Simplify x-axis labels based on time span
        if year_span >= 8:
            step = 12  # Show yearly
        elif year_span >= 5:
            step = 6   # Show every 6 months
        elif year_span >= 3:
            step = 3   # Show quarterly
        elif year_span >= 2:
            step = 2   # Show every 2 months
        else:
            step = 1   # Show monthly (1 year or less)

        # Format month labels as MMM YY (e.g., Jan 23)
        month_names = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
                       'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
        month_labels = []
        for m in months:
            year = m[:4]
            month_num = int(m[5:7])
            month_labels.append(f"{month_names[month_num-1]} {year[2:]}")

        # Only show labels at step intervals
        visible_ticks = list(range(0, len(months), step))
        ax.set_xticks(visible_ticks)
        ax.set_xticklabels([month_labels[i] for i in visible_ticks], rotation=45, ha='right', fontsize=9)

        ax.set_ylabel('Incidents', color='white', fontsize=10)
        ax.set_title('Risk Incidents Over Time', color='white', fontsize=12)
        ax.tick_params(colors='white', labelsize=9)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['bottom'].set_color('#444')
        ax.spines['left'].set_color('#444')

        # Legend outside chart on the right
        ax.legend(loc='upper right', fontsize=8, facecolor='#2a2a3e', edgecolor='#444',
                 labelcolor='white', framealpha=0.9)

        # Pin Y-axis to global max so filtered views stay proportional
        if y_max is not None and y_max > 0:
            ax.set_ylim(0, y_max + 1)

        plt.tight_layout()

        # Calculate bar positions for interactivity
        fig_height_px = fig_height * dpi
        fig_width_px = fig_width * dpi

        bar_positions = []  # [(month_key, category, x_center_px, y_top_px, width_px, height_px, count), ...]

        for cat, bars in all_bars:
            for month_idx, bar in enumerate(bars):
                if bar.get_height() > 0:
                    # Get bar corners in display coordinates
                    x_data = bar.get_x()
                    w_data = bar.get_width()
                    h_data = bar.get_height()

                    # Transform to pixel coordinates
                    x_left_px = ax.transData.transform((x_data, 0))[0]
                    x_right_px = ax.transData.transform((x_data + w_data, 0))[0]
                    y_bottom_px = ax.transData.transform((0, 0))[1]
                    y_top_px = ax.transData.transform((0, h_data))[1]

                    # Convert to Qt coordinates (flip Y)
                    x_center_px = (x_left_px + x_right_px) / 2
                    width_px = x_right_px - x_left_px
                    height_px = y_top_px - y_bottom_px
                    y_qt = fig_height_px - y_top_px

                    month_key = months[month_idx]
                    count = int(h_data)

                    bar_positions.append({
                        'month': month_key,
                        'month_label': month_labels[month_idx],
                        'category': cat,
                        'x': x_center_px,
                        'y': y_qt,
                        'width': width_px,
                        'height': height_px,
                        'count': count
                    })

        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, facecolor='#1a1a2e', edgecolor='none')
        plt.close(fig)
        buf.seek(0)

        timeline_info = {
            'months': months,
            'month_labels': month_labels,
            'categories': top_cats,
            'monthly_counts': dict(monthly_counts),
            'bar_positions': bar_positions
        }

        return buf.read(), timeline_info

    except Exception as e:
        print(f"Timeline chart error: {e}")
        import traceback
        traceback.print_exc()
        return None, {}


def create_pie_chart_with_legend(categories_data: Dict) -> bytes:
    """Create a pie chart with legend for Word export."""
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        # Filter categories with counts > 0
        data = [(name, info["count"], info["color"], info.get("icon", ""))
                for name, info in categories_data.items()
                if info["count"] > 0]

        if not data:
            return None

        # Sort by count
        data.sort(key=lambda x: x[1], reverse=True)

        names = [d[0] for d in data]
        counts = [d[1] for d in data]
        colors = [d[2] for d in data]
        icons = [d[3] for d in data]
        total = sum(counts)

        # Create figure with space for legend
        fig, ax = plt.subplots(figsize=(8, 5))
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')

        # Create pie chart
        wedges, texts, autotexts = ax.pie(
            counts,
            labels=None,
            colors=colors,
            autopct=lambda pct: f'{pct:.0f}%' if pct > 5 else '',
            startangle=90,
            pctdistance=0.75,
            wedgeprops=dict(width=0.6, edgecolor='white', linewidth=2)
        )

        # Style percentage text
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontsize(10)
            autotext.set_fontweight('bold')

        # Add legend
        legend_labels = [f"{icons[i]} {names[i]} ({counts[i]})" for i in range(len(names))]
        ax.legend(wedges, legend_labels, title="Risk Categories", loc="center left",
                  bbox_to_anchor=(1, 0, 0.5, 1), fontsize=9)

        plt.tight_layout()

        # Save to bytes
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, facecolor='white', edgecolor='none',
                    bbox_inches='tight', pad_inches=0.1)
        plt.close(fig)
        buf.seek(0)
        return buf.read()

    except Exception as e:
        print(f"Pie chart with legend error: {e}")
        return None


def export_risk_to_word(results: Dict, output_path: str) -> str:
    """Export risk analysis to a Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('Risk Overview Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Summary
        total = results["total_notes"]
        with_incidents = results["notes_with_incidents"]
        pct = (with_incidents / total * 100) if total > 0 else 0

        summary = doc.add_paragraph()
        summary.add_run(f"Total notes analyzed: {total}\n").bold = True
        summary.add_run(f"Notes with risk indicators: {with_incidents} ({pct:.1f}%)\n")

        # Add category chart with legend
        chart_bytes = create_pie_chart_with_legend(results["categories"])
        if chart_bytes:
            chart_stream = io.BytesIO(chart_bytes)
            doc.add_heading('Risk Categories Overview', level=1)
            doc.add_picture(chart_stream, width=Inches(5.5))
            doc.add_paragraph()

        # Risk Evidence by Category
        doc.add_heading('Risk Evidence by Category', level=1)

        sorted_cats = sorted(
            results["categories"].items(),
            key=lambda x: x[1]["count"],
            reverse=True
        )

        for cat_name, cat_data in sorted_cats:
            if cat_data["count"] == 0:
                continue

            # Category heading
            doc.add_heading(f"{cat_data['icon']} {cat_name}", level=2)

            # Deduplicate incidents by date - keep most relevant per date
            # Priority: higher severity, then longer matched text
            severity_rank = {"high": 3, "medium": 2, "low": 1}
            incidents_by_date = {}
            for inc in cat_data["incidents"]:
                date_key = inc["date"].strftime("%Y-%m-%d") if inc["date"] else "unknown"
                inc_severity = severity_rank.get(inc.get("severity", "low"), 1)
                inc_match_len = len(inc.get("matched", ""))

                if date_key not in incidents_by_date:
                    incidents_by_date[date_key] = inc
                else:
                    # Compare and keep more relevant
                    existing = incidents_by_date[date_key]
                    existing_severity = severity_rank.get(existing.get("severity", "low"), 1)
                    existing_match_len = len(existing.get("matched", ""))

                    # Prefer higher severity, then longer match
                    if inc_severity > existing_severity or \
                       (inc_severity == existing_severity and inc_match_len > existing_match_len):
                        incidents_by_date[date_key] = inc

            # Sort deduplicated incidents by date
            deduped_incidents = sorted(incidents_by_date.values(),
                                       key=lambda x: x["date"] if x["date"] else datetime.min,
                                       reverse=True)

            # List incidents - date and full sentence containing the match
            from docx.enum.text import WD_COLOR_INDEX
            for inc in deduped_incidents:
                date_str = inc["date"].strftime("%d %b %Y") if inc["date"] else "Unknown"

                # Get the full text and find the sentence containing the match
                full_text = inc.get("full_text", "")
                matched = inc.get("matched", "")

                if matched and full_text:
                    # Find the sentence containing the matched term
                    match_pos = full_text.lower().find(matched.lower())
                    if match_pos != -1:
                        # Find sentence boundaries (., !, ?, or newline)
                        # Look backwards for start
                        start = match_pos
                        while start > 0 and full_text[start-1] not in '.!?\n':
                            start -= 1
                        # Look forwards for end
                        end = match_pos + len(matched)
                        while end < len(full_text) and full_text[end] not in '.!?\n':
                            end += 1
                        # Include the punctuation
                        if end < len(full_text) and full_text[end] in '.!?':
                            end += 1
                        sentence = full_text[start:end].strip()
                        # Calculate position of match within extracted sentence
                        match_in_sentence = sentence.lower().find(matched.lower())
                    else:
                        sentence = matched
                        match_in_sentence = 0
                else:
                    sentence = matched or full_text[:300]
                    match_in_sentence = 0 if matched else -1

                # Clean up and limit length
                sentence = sentence.strip()
                if len(sentence) > 400:
                    sentence = sentence[:400].rsplit(' ', 1)[0] + "..."
                    # Recalculate match position after truncation
                    if matched:
                        match_in_sentence = sentence.lower().find(matched.lower())

                p = doc.add_paragraph()
                p.add_run(f"{date_str}: ").bold = True

                # Add sentence with highlighted matched term
                if matched and match_in_sentence >= 0:
                    # Text before match
                    if match_in_sentence > 0:
                        p.add_run(sentence[:match_in_sentence])
                    # Highlighted match
                    highlight_run = p.add_run(sentence[match_in_sentence:match_in_sentence + len(matched)])
                    highlight_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    # Text after match
                    if match_in_sentence + len(matched) < len(sentence):
                        p.add_run(sentence[match_in_sentence + len(matched):])
                else:
                    p.add_run(sentence)

                p.paragraph_format.space_after = Pt(4)

            doc.add_paragraph()

        doc.save(output_path)
        return output_path

    except Exception as e:
        print(f"Word export error: {e}")
        return None


def create_pie_chart(categories_data: Dict, embedded: bool = False) -> tuple:
    """Create a pie chart using matplotlib.

    Returns:
        tuple: (image_bytes, wedge_info_dict) where wedge_info_dict contains
               {category_name: {"count": N, "color": "#xxx", "start_angle": deg, "end_angle": deg, "center": (x,y), "radius": r}, ...}
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np

        # Filter categories with counts > 0
        data = [(name, info["count"], info["color"], info.get("icon", ""))
                for name, info in categories_data.items()
                if info["count"] > 0]

        if not data:
            return None, {}

        # Sort by count
        data.sort(key=lambda x: x[1], reverse=True)

        names = [d[0] for d in data]
        counts = [d[1] for d in data]
        colors = [d[2] for d in data]
        icons = [d[3] for d in data]
        total = sum(counts)

        # Create square figure (no legend - we'll add Qt legend separately)
        fig_size = 5
        fig, ax = plt.subplots(figsize=(fig_size, fig_size))

        # Background color based on mode
        bg_color = 'white' if embedded else '#1a1a2e'
        fig.patch.set_facecolor(bg_color)
        ax.set_facecolor(bg_color)

        # Create pie chart
        wedges, texts, autotexts = ax.pie(
            counts,
            labels=None,
            colors=colors,
            autopct=lambda pct: f'{pct:.0f}%' if pct > 5 else '',
            startangle=90,
            pctdistance=0.75,
            wedgeprops=dict(width=0.6, edgecolor=bg_color, linewidth=2)
        )

        # Style percentage text
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontsize(12)
            autotext.set_fontweight('bold')

        plt.tight_layout()

        # Build category info for the interactive legend
        category_info = {}
        total_count = total
        for i, (name, count) in enumerate(zip(names, counts)):
            category_info[name] = {
                "count": count,
                "color": colors[i],
                "icon": icons[i],
                "percentage": (count / total_count) * 100
            }

        # Save to bytes
        dpi = 100
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, facecolor=bg_color, edgecolor='none',
                    bbox_inches='tight', pad_inches=0.05)
        plt.close(fig)
        buf.seek(0)
        return buf.read(), category_info

    except Exception as e:
        print(f"Pie chart creation error: {e}")
        import traceback
        traceback.print_exc()
        return None, {}


# ============================================================
# PIE CHART WITH INTERACTIVE LEGEND
# ============================================================

class PieChartWithLegend(QWidget):
    """Pie chart with interactive legend - tooltips and click dropdown on legend items."""

    incidentSelected = Signal(object)  # Emits incident dict when selected

    def __init__(self, parent=None, embedded=True):
        super().__init__(parent)
        self.embedded = embedded
        self.categories_data = {}
        self.category_info = {}
        self._build_ui()

    def _build_ui(self):
        """Build the legend + chart layout (legend on left)."""
        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Legend container (on left - first thing seen)
        self.legend_widget = QWidget()
        self.legend_layout = QVBoxLayout(self.legend_widget)
        self.legend_layout.setContentsMargins(8, 8, 8, 8)
        self.legend_layout.setSpacing(6)
        self.legend_layout.setAlignment(Qt.AlignTop)
        layout.addWidget(self.legend_widget, stretch=0)

        # Pie chart image (on right)
        self.chart_label = QLabel()
        self.chart_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(self.chart_label, stretch=1)

    def set_data(self, image_bytes: bytes, category_info: dict, categories_data: dict):
        """Set the chart image and data."""
        self.category_info = category_info
        self.categories_data = categories_data

        # Set chart image
        if image_bytes:
            pixmap = QPixmap()
            pixmap.loadFromData(image_bytes)
            self.chart_label.setPixmap(pixmap)
            self.chart_label.adjustSize()

        # Build interactive legend
        self._build_legend()

    def _build_legend(self):
        """Create interactive legend labels."""
        # Clear existing legend
        while self.legend_layout.count():
            item = self.legend_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Add legend items sorted by count
        sorted_cats = sorted(self.category_info.items(), key=lambda x: x[1]["count"], reverse=True)

        for name, info in sorted_cats:
            cat_data = self.categories_data.get(name, {})
            color = info.get("color", "#888")
            count = info.get("count", 0)
            percentage = info.get("percentage", 0)
            icon = cat_data.get("icon", "")

            # Create legend item button with percentage
            btn = QPushButton(f"● {name}: {count} ({percentage:.0f}%)")
            btn.setCursor(Qt.PointingHandCursor)
            btn.setToolTip(self._build_tooltip(name, info, cat_data))

            # Style with category color
            if self.embedded:
                btn.setStyleSheet(f"""
                    QPushButton {{
                        background: transparent;
                        border: none;
                        color: {color};
                        font-size: 13px;
                        font-weight: 600;
                        text-align: left;
                        padding: 6px 10px;
                        border-radius: 6px;
                    }}
                    QPushButton:hover {{
                        background: rgba(0,0,0,0.08);
                    }}
                """)
            else:
                btn.setStyleSheet(f"""
                    QPushButton {{
                        background: transparent;
                        border: none;
                        color: {color};
                        font-size: 13px;
                        font-weight: 600;
                        text-align: left;
                        padding: 6px 10px;
                        border-radius: 6px;
                    }}
                    QPushButton:hover {{
                        background: rgba(255,255,255,0.15);
                    }}
                """)

            btn.clicked.connect(lambda checked, n=name: self._show_category_menu(n))
            self.legend_layout.addWidget(btn)

        self.legend_layout.addStretch()

    def _build_tooltip(self, name: str, info: dict, cat_data: dict) -> str:
        """Build HTML tooltip for a category."""
        tooltip = f"<b>{cat_data.get('icon', '')} {name}</b><br>"
        tooltip += f"<b>{info['count']}</b> incidents ({info['percentage']:.1f}%)<br>"

        # Show subcategory breakdown if available
        subcategories = cat_data.get("subcategories", {})
        if subcategories:
            subcat_counts = [(n, d["count"], d.get("severity", "medium"))
                             for n, d in subcategories.items() if d["count"] > 0]
            subcat_counts.sort(key=lambda x: x[1], reverse=True)

            if subcat_counts:
                tooltip += "<br><b>Types:</b><br>"
                for subcat_name, subcat_count, severity in subcat_counts[:5]:
                    if severity == "high":
                        sev_color = "#ef4444"
                    elif severity == "medium":
                        sev_color = "#f59e0b"
                    else:
                        sev_color = "#22c55e"
                    tooltip += f"<span style='color: {sev_color};'>●</span> {subcat_name}: {subcat_count}<br>"
                if len(subcat_counts) > 5:
                    tooltip += f"<i>+{len(subcat_counts) - 5} more</i>"

        return tooltip

    def _show_category_menu(self, name: str):
        """Show scrollable popup with incidents for a category."""
        from PySide6.QtWidgets import QListWidget, QListWidgetItem

        if name not in self.categories_data:
            return

        cat_data = self.categories_data[name]
        incidents = cat_data.get("incidents", [])

        if not incidents:
            return

        # Close any existing popup
        if hasattr(self, '_active_popup') and self._active_popup:
            try:
                self._active_popup.close()
                self._active_popup.deleteLater()
            except:
                pass
            self._active_popup = None

        # Sort incidents by date (most recent first)
        sorted_incidents = sorted(incidents, key=lambda x: x.get('date') or datetime.min, reverse=True)

        # Create popup as a child widget (not a separate window)
        popup = QFrame(self.window())
        popup.setWindowFlags(Qt.Popup)
        popup.setFixedSize(400, 400)
        popup.setStyleSheet("""
            QFrame {
                background-color: rgb(45,45,45);
                border: 2px solid #888;
                border-radius: 8px;
            }
        """)

        layout = QVBoxLayout(popup)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        # Header with close button
        header_layout = QHBoxLayout()
        header = QLabel(f"{cat_data.get('icon', '')} {name} ({len(incidents)} incidents)")
        header.setStyleSheet("font-size: 14px; font-weight: bold; color: white; padding: 6px; background: transparent;")
        header_layout.addWidget(header)
        header_layout.addStretch()

        close_btn = QPushButton("✕")
        close_btn.setFixedSize(24, 24)
        close_btn.setCursor(Qt.PointingHandCursor)
        close_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.2);
                border: none;
                border-radius: 12px;
                color: white;
                font-size: 14px;
            }
            QPushButton:hover {
                background: rgba(255,100,100,0.6);
            }
        """)
        close_btn.clicked.connect(popup.close)
        header_layout.addWidget(close_btn)
        layout.addLayout(header_layout)

        # Scrollable list
        list_widget = QListWidget()
        list_widget.setStyleSheet("""
            QListWidget {
                background-color: transparent;
                border: none;
                color: white;
                font-size: 12px;
            }
            QListWidget::item {
                padding: 8px 10px;
                border-radius: 4px;
                margin: 2px 0;
            }
            QListWidget::item:hover {
                background-color: rgba(70,130,180,0.5);
            }
            QListWidget::item:selected {
                background-color: rgba(70,130,180,0.7);
            }
            QScrollBar:vertical {
                background: rgba(255,255,255,0.1);
                width: 10px;
                border-radius: 5px;
            }
            QScrollBar::handle:vertical {
                background: rgba(255,255,255,0.3);
                border-radius: 5px;
                min-height: 30px;
            }
        """)

        # Add all incidents to the list
        for inc in sorted_incidents:
            date_str = inc.get('date').strftime('%d %b %Y') if inc.get('date') else 'Unknown'
            subcat = inc.get('subcategory', '')
            text = inc.get('text', '')[:80] + ('...' if len(inc.get('text', '')) > 80 else '')

            item = QListWidgetItem(f"{date_str} - {subcat}")
            item.setToolTip(text)
            item.setData(Qt.UserRole, inc)
            list_widget.addItem(item)

        # Handle item click
        def on_item_clicked(item):
            inc = item.data(Qt.UserRole)
            if inc:
                self.incidentSelected.emit(inc)
            popup.close()

        list_widget.itemClicked.connect(on_item_clicked)
        layout.addWidget(list_widget)

        # Store reference and show
        self._active_popup = popup

        # Position popup near the button that was clicked, within main window
        popup.move(QCursor.pos())
        popup.show()


# ============================================================
# INTERACTIVE BAR CHART - Tooltips and Click-to-Scroll
# ============================================================

class InteractiveBarChart(QLabel):
    """Bar chart with hover tooltips and click-to-scroll functionality."""

    incidentSelected = Signal(object)  # Emits incident dict when selected
    exportRequested = Signal(str)      # Emits category name to export

    def __init__(self, parent=None):
        super().__init__(parent)
        self.bar_info = []  # [(name, count, y_center, height), ...]
        self.categories_data = {}  # Full category data with incidents
        self.setMouseTracking(True)
        self.setCursor(Qt.PointingHandCursor)

    def set_chart(self, image_bytes: bytes, bar_info: list, categories_data: dict = None):
        """Set the chart image, bar position info, and category data."""
        self.bar_info = bar_info
        self.categories_data = categories_data or {}
        if image_bytes:
            pixmap = QPixmap()
            pixmap.loadFromData(image_bytes)
            self.setPixmap(pixmap)

    def _get_bar_at_pos(self, pos):
        """Return bar info if position is over a bar."""
        y = pos.y()
        for name, count, y_center, height in self.bar_info:
            if abs(y - y_center) <= height / 2 + 5:  # +5px tolerance
                return name, count
        return None, None

    def mouseMoveEvent(self, event):
        """Show tooltip with source data on hover."""
        name, count = self._get_bar_at_pos(event.pos())
        if name and name in self.categories_data:
            cat_data = self.categories_data[name]
            incidents = cat_data.get("incidents", [])
            subcategories = cat_data.get("subcategories", {})

            # Build tooltip with incident snippets
            tooltip = f"<b>{name}</b> - {count} incident{'s' if count != 1 else ''}<br>"

            # Show subcategory breakdown if available
            if subcategories:
                subcat_counts = [(n, d["count"], d.get("severity", "medium"))
                                 for n, d in subcategories.items() if d["count"] > 0]
                subcat_counts.sort(key=lambda x: x[1], reverse=True)

                if subcat_counts:
                    tooltip += "<br><b>Breakdown:</b><br>"
                    for subcat_name, subcat_count, severity in subcat_counts[:4]:
                        # Color code by severity
                        if severity == "high":
                            sev_color = "#ef4444"
                        elif severity == "medium":
                            sev_color = "#f59e0b"
                        else:
                            sev_color = "#22c55e"
                        tooltip += f"<span style='color: {sev_color};'>●</span> {subcat_name}: {subcat_count}<br>"
                    if len(subcat_counts) > 4:
                        tooltip += f"<i>+{len(subcat_counts) - 4} more types</i><br>"

            tooltip += "<br>"

            # Show up to 3 recent incidents
            for inc in incidents[:3]:
                date_str = inc["date"].strftime("%d %b %Y") if inc.get("date") else "Unknown"
                subcat = inc.get("subcategory", "")
                severity = inc.get("severity", "")

                # Get snippet around matched text
                matched = inc.get("matched", "")
                full_text = inc.get("full_text", "")
                if matched and full_text:
                    idx = full_text.lower().find(matched.lower())
                    if idx >= 0:
                        start = max(0, idx - 30)
                        end = min(len(full_text), idx + len(matched) + 30)
                        snippet = full_text[start:end]
                        if start > 0:
                            snippet = "..." + snippet
                        if end < len(full_text):
                            snippet = snippet + "..."
                    else:
                        snippet = full_text[:60] + "..."
                else:
                    snippet = full_text[:60] + "..." if len(full_text) > 60 else full_text

                # Add severity indicator
                sev_badge = ""
                if severity == "high":
                    sev_badge = "<span style='color: #ef4444;'>[HIGH]</span> "
                elif severity == "medium":
                    sev_badge = "<span style='color: #f59e0b;'>[MED]</span> "

                tooltip += f"<b>{date_str}:</b> {sev_badge}{snippet}<br>"

            if len(incidents) > 3:
                tooltip += f"<i>... and {len(incidents) - 3} more</i><br>"

            tooltip += "<br><i>Click to view incidents</i>"

            # Wrap in HTML with explicit styling for Windows
            tooltip = f"<div style='background-color: #fffbe6; color: #000000; padding: 4px;'>{tooltip}</div>"
            QToolTip.showText(event.globalPosition().toPoint(), tooltip, self)
        else:
            QToolTip.hideText()
        super().mouseMoveEvent(event)

    def mousePressEvent(self, event):
        """Show menu of incidents when bar is clicked."""
        if event.button() == Qt.LeftButton:
            # Hide tooltip before showing menu (Windows fix)
            QToolTip.hideText()

            name, count = self._get_bar_at_pos(event.pos())
            if name and name in self.categories_data:
                cat_data = self.categories_data[name]
                incidents = cat_data.get("incidents", [])

                if not incidents:
                    return

                # Show context menu with incidents
                menu = QMenu(self)
                menu.setAttribute(Qt.WA_DeleteOnClose)  # Ensure cleanup on Windows
                menu.setStyleSheet("""
                    QMenu {
                        background-color: rgba(40,40,40,0.95);
                        color: white;
                        border: 1px solid #555;
                        border-radius: 6px;
                        padding: 4px;
                        min-width: 300px;
                    }
                    QMenu::item {
                        padding: 8px 16px;
                        border-radius: 4px;
                    }
                    QMenu::item:selected {
                        background-color: rgba(70,130,180,0.7);
                    }
                    QMenu::separator {
                        height: 1px;
                        background: #555;
                        margin: 4px 8px;
                    }
                """)

                # Add incidents as menu items (limit to 15)
                for inc in incidents[:15]:
                    date_str = inc["date"].strftime("%d %b %Y") if inc.get("date") else "Unknown"
                    matched = inc.get("matched", "")[:25]
                    severity = inc.get("severity", "")
                    subcat = inc.get("subcategory", "")

                    # Add severity indicator
                    sev_icon = ""
                    if severity == "high":
                        sev_icon = "[!] "
                    elif severity == "medium":
                        sev_icon = "[*] "

                    action = menu.addAction(f"{sev_icon}{date_str}: {matched}...")
                    action.setData(inc)

                if len(incidents) > 15:
                    menu.addSeparator()
                    menu.addAction(f"... {len(incidents) - 15} more incidents")

                menu.addSeparator()
                export_action = menu.addAction(f"Export all {name} to Word")

                action = menu.exec(event.globalPosition().toPoint())
                if action:
                    if action == export_action:
                        self.exportRequested.emit(name)
                    elif action.data():
                        self.incidentSelected.emit(action.data())

        super().mousePressEvent(event)


# ============================================================
# INTERACTIVE TIMELINE CHART - Tooltips for time-based data
# ============================================================

class InteractiveTimelineChart(QLabel):
    """Timeline chart with hover tooltips, click-to-scroll, and visual zoom."""

    incidentSelected = Signal(object)  # Emits incident dict when selected
    zoomChanged = Signal(int)  # Emits zoom percentage when changed

    def __init__(self, parent=None):
        super().__init__(parent)
        self.timeline_info = {}
        self.categories_data = {}
        self.original_pixmap = None
        self.zoom_factor = 1.0
        self.min_zoom = 1.0
        self.max_zoom = 4.0
        self.setMouseTracking(True)
        self.setCursor(Qt.PointingHandCursor)

    def zoom_in(self):
        """Zoom in by 25%."""
        if self.original_pixmap is None:
            return
        self.zoom_factor = min(self.max_zoom, self.zoom_factor + 0.25)
        self._apply_zoom()

    def zoom_out(self):
        """Zoom out by 25%."""
        if self.original_pixmap is None:
            return
        self.zoom_factor = max(self.min_zoom, self.zoom_factor - 0.25)
        self._apply_zoom()

    def _apply_zoom(self):
        """Apply current zoom factor to the pixmap."""
        if self.original_pixmap is None:
            return

        new_width = int(self.original_pixmap.width() * self.zoom_factor)
        new_height = int(self.original_pixmap.height() * self.zoom_factor)

        scaled = self.original_pixmap.scaled(
            new_width, new_height,
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation
        )
        self.setPixmap(scaled)
        self.adjustSize()  # Resize label to fit scaled pixmap
        self.zoomChanged.emit(int(self.zoom_factor * 100))

    def reset_zoom(self):
        """Reset zoom to 100%."""
        self.zoom_factor = 1.0
        if self.original_pixmap:
            self.setPixmap(self.original_pixmap)
            self.adjustSize()
        self.zoomChanged.emit(100)

    def set_chart(self, image_bytes: bytes, timeline_info: dict, categories_data: dict = None):
        """Set the chart image and timeline info for tooltips."""
        self.timeline_info = timeline_info
        self.categories_data = categories_data or {}
        self.zoom_factor = 1.0  # Reset zoom when new chart loaded
        if image_bytes:
            self.original_pixmap = QPixmap()
            self.original_pixmap.loadFromData(image_bytes)
            self.setPixmap(self.original_pixmap)
            self.adjustSize()  # Resize label to fit pixmap
            self.zoomChanged.emit(100)

    def _get_bar_at_pos(self, pos):
        """Return bar info if position is over a bar (adjusted for zoom)."""
        if not self.timeline_info.get('bar_positions'):
            return None

        x, y = pos.x(), pos.y()

        # Scale positions by zoom factor
        zf = self.zoom_factor

        for bar in self.timeline_info['bar_positions']:
            # Check if point is within bar bounds (scaled by zoom)
            bar_left = (bar['x'] - bar['width'] / 2) * zf
            bar_right = (bar['x'] + bar['width'] / 2) * zf
            bar_top = bar['y'] * zf
            bar_bottom = (bar['y'] + bar['height']) * zf

            if bar_left <= x <= bar_right and bar_top <= y <= bar_bottom:
                return bar

        return None

    def mouseMoveEvent(self, event):
        """Show tooltip with incident details on hover."""
        bar = self._get_bar_at_pos(event.pos())

        if bar:
            month_label = bar['month_label']
            category = bar['category']
            count = bar['count']

            # Build tooltip with explicit styling for Windows
            tooltip = "<div style='background-color: #fffbe6; color: #000000; padding: 4px;'>"
            tooltip += f"<b>{month_label}</b><br>"
            tooltip += f"<b>{category}:</b> {count} incident{'s' if count != 1 else ''}<br>"

            # Try to get incident snippets for this month/category
            if self.categories_data and category in self.categories_data:
                incidents = self.categories_data[category].get('incidents', [])
                month_key = bar['month']

                # Filter incidents for this month
                month_incidents = []
                for inc in incidents:
                    if inc.get('date'):
                        inc_month = inc['date'].strftime('%Y-%m')
                        if inc_month == month_key:
                            month_incidents.append(inc)

                if month_incidents:
                    tooltip += "<br>"
                    for inc in month_incidents[:3]:
                        date_str = inc['date'].strftime('%d %b') if inc.get('date') else ''
                        matched = inc.get('matched', '')[:25]
                        severity = inc.get('severity', '')
                        subcat = inc.get('subcategory', '')

                        # Severity indicator
                        sev_badge = ""
                        if severity == "high":
                            sev_badge = "<span style='color: #dc2626;'>[HIGH]</span> "
                        elif severity == "medium":
                            sev_badge = "<span style='color: #d97706;'>[MED]</span> "

                        tooltip += f"• {date_str}: {sev_badge}{matched}...<br>"

                    if len(month_incidents) > 3:
                        tooltip += f"<i>+{len(month_incidents) - 3} more</i>"

            tooltip += "</div>"
            QToolTip.showText(event.globalPosition().toPoint(), tooltip, self)
        else:
            QToolTip.hideText()

        super().mouseMoveEvent(event)

    def mousePressEvent(self, event):
        """Show menu of incidents when bar is clicked."""
        if event.button() == Qt.LeftButton:
            # Hide tooltip before showing menu (Windows fix)
            QToolTip.hideText()

            # Use fresh cursor position for Windows compatibility
            local_pos = self.mapFromGlobal(QCursor.pos())
            bar = self._get_bar_at_pos(local_pos)

            if bar and self.categories_data:
                category = bar['category']
                month_key = bar['month']
                month_label = bar['month_label']

                if category in self.categories_data:
                    incidents = self.categories_data[category].get('incidents', [])

                    # Filter incidents for this month
                    month_incidents = []
                    for inc in incidents:
                        if inc.get('date'):
                            inc_month = inc['date'].strftime('%Y-%m')
                            if inc_month == month_key:
                                month_incidents.append(inc)

                    if month_incidents:
                        # Show context menu with incidents
                        menu = QMenu(self)
                        menu.setAttribute(Qt.WA_DeleteOnClose)  # Ensure cleanup on Windows
                        menu.setStyleSheet("""
                            QMenu {
                                background-color: rgba(40,40,40,0.95);
                                color: white;
                                border: 1px solid #555;
                                border-radius: 6px;
                                padding: 4px;
                                min-width: 280px;
                            }
                            QMenu::item {
                                padding: 8px 16px;
                                border-radius: 4px;
                            }
                            QMenu::item:selected {
                                background-color: rgba(70,130,180,0.7);
                            }
                            QMenu::separator {
                                height: 1px;
                                background: #555;
                                margin: 4px 8px;
                            }
                        """)

                        # Header
                        header = menu.addAction(f"{month_label} - {category}")
                        header.setEnabled(False)
                        menu.addSeparator()

                        # Add incidents as menu items
                        for inc in month_incidents[:15]:
                            date_str = inc['date'].strftime('%d %b %Y') if inc.get('date') else 'Unknown'
                            matched = inc.get('matched', '')[:25]
                            severity = inc.get('severity', '')

                            # Severity indicator
                            sev_icon = ""
                            if severity == "high":
                                sev_icon = "[!] "
                            elif severity == "medium":
                                sev_icon = "[*] "

                            action = menu.addAction(f"{sev_icon}{date_str}: {matched}...")
                            action.setData(inc)

                        if len(month_incidents) > 15:
                            menu.addSeparator()
                            more = menu.addAction(f"... {len(month_incidents) - 15} more")
                            more.setEnabled(False)

                        action = menu.exec(QCursor.pos())
                        if action and action.data():
                            self.incidentSelected.emit(action.data())

                        # Force refresh after menu closes (Windows fix)
                        self.update()

        event.accept()


# ============================================================
# INTERACTIVE RISK LEVEL TIMELINE - Horizontal bar chart with tooltips
# ============================================================

class InteractiveRiskLevelTimeline(QLabel):
    """Interactive horizontal risk level timeline with tooltips and click-to-scroll."""

    incidentSelected = Signal(object)

    def __init__(self):
        super().__init__()
        self.setMouseTracking(True)
        self.setCursor(Qt.PointingHandCursor)
        self.timeline_info = {}
        self.categories_data = {}

    def set_chart(self, image_bytes: bytes, timeline_info: dict, categories_data: dict):
        """Set the chart image and metadata."""
        self.timeline_info = timeline_info
        self.categories_data = categories_data
        if image_bytes:
            pixmap = QPixmap()
            pixmap.loadFromData(image_bytes)
            self.setPixmap(pixmap)

    def _get_bar_at_pos(self, pos):
        """Return bar info if position is over a bar."""
        if not self.timeline_info.get('bar_positions'):
            return None

        x, y = pos.x(), pos.y()

        for bar in self.timeline_info['bar_positions']:
            if bar['x_left'] <= x <= bar['x_right'] and bar['y_top'] <= y <= bar['y_bottom']:
                return bar

        return None

    def mouseMoveEvent(self, event):
        """Show tooltip with month risk details on hover."""
        bar = self._get_bar_at_pos(event.pos())

        if bar:
            month = bar['month']
            total = bar['total']
            level_name = bar['level_name']
            top_cats = bar.get('top_cats', [])

            # Format month name
            try:
                dt = datetime.strptime(month, "%Y-%m")
                month_label = dt.strftime("%B %Y")
            except:
                month_label = month

            # Build tooltip
            tooltip = f"<b>{month_label}</b><br>"
            tooltip += f"Risk Level: <b>{level_name}</b><br>"
            tooltip += f"Total Incidents: <b>{total}</b><br>"

            if top_cats:
                tooltip += "<br><b>Categories:</b><br>"
                for cat, count in top_cats:
                    tooltip += f"• {cat}: {count}<br>"

            if total > 0:
                tooltip += "<br><i>Click to view incidents</i>"

            # Wrap in HTML with explicit styling for Windows
            tooltip = f"<div style='background-color: #fffbe6; color: #000000; padding: 4px;'>{tooltip}</div>"
            QToolTip.showText(event.globalPosition().toPoint(), tooltip, self)
        else:
            QToolTip.hideText()

        super().mouseMoveEvent(event)

    def mousePressEvent(self, event):
        """Show menu of incidents when bar is clicked."""
        if event.button() == Qt.LeftButton:
            # Hide tooltip before showing menu (Windows fix)
            QToolTip.hideText()

            # Use fresh cursor position for Windows compatibility
            local_pos = self.mapFromGlobal(QCursor.pos())
            bar = self._get_bar_at_pos(local_pos)

            if bar and bar['total'] > 0:
                month = bar['month']

                # Format month name
                try:
                    dt = datetime.strptime(month, "%Y-%m")
                    month_label = dt.strftime("%B %Y")
                except:
                    month_label = month

                # Gather incidents for this month from all categories
                month_incidents = []
                for cat_name, cat_data in self.categories_data.items():
                    for inc in cat_data.get('incidents', []):
                        if inc.get('date'):
                            inc_month = inc['date'].strftime('%Y-%m')
                            if inc_month == month:
                                inc_copy = dict(inc)
                                inc_copy['category'] = cat_name
                                month_incidents.append(inc_copy)

                if month_incidents:
                    # Sort by date
                    month_incidents.sort(key=lambda x: x.get('date') or datetime.min)

                    # Show context menu
                    menu = QMenu(self)
                    menu.setAttribute(Qt.WA_DeleteOnClose)  # Ensure cleanup on Windows
                    menu.setStyleSheet("""
                        QMenu {
                            background-color: rgba(40,40,40,0.95);
                            color: white;
                            border: 1px solid #555;
                            border-radius: 6px;
                            padding: 4px;
                            min-width: 320px;
                        }
                        QMenu::item {
                            padding: 8px 16px;
                            border-radius: 4px;
                        }
                        QMenu::item:selected {
                            background-color: rgba(70,130,180,0.7);
                        }
                        QMenu::separator {
                            height: 1px;
                            background: #555;
                            margin: 4px 8px;
                        }
                    """)

                    # Header
                    header = menu.addAction(f"{month_label} - {bar['level_name']} ({bar['total']} incidents)")
                    header.setEnabled(False)
                    menu.addSeparator()

                    # Add incidents
                    for inc in month_incidents[:15]:
                        date_str = inc['date'].strftime('%d %b') if inc.get('date') else ''
                        category = inc.get('category', '')
                        matched = inc.get('matched', '')[:20]
                        severity = inc.get('severity', '')

                        # Severity indicator
                        sev_icon = ""
                        if severity == "high":
                            sev_icon = "[!] "
                        elif severity == "medium":
                            sev_icon = "[*] "

                        action = menu.addAction(f"{sev_icon}{date_str} [{category}]: {matched}...")
                        action.setData(inc)

                    if len(month_incidents) > 15:
                        menu.addSeparator()
                        more = menu.addAction(f"... {len(month_incidents) - 15} more")
                        more.setEnabled(False)

                    action = menu.exec(QCursor.pos())
                    if action and action.data():
                        self.incidentSelected.emit(action.data())

                    # Force refresh after menu closes (Windows fix)
                    self.update()

        event.accept()


# ============================================================
# RISK OVERVIEW PANEL - Matches Patient History Panel Style
# ============================================================

class RiskOverviewPanel(QWidget):
    """Floating panel displaying risk analysis from notes."""

    closed = Signal()

    def __init__(self, notes: List[Dict], parent=None, notes_panel=None, embedded=False):
        super().__init__(parent)
        self.notes = notes
        self.notes_panel = notes_panel  # Reference to left notes panel for scroll-to
        self.embedded = embedded

        self._drag_offset = QPoint()
        self._dragging = False

        # Window settings - only for floating mode
        if not embedded:
            self.setWindowFlags(
                Qt.FramelessWindowHint |
                Qt.SubWindow |
                Qt.WindowStaysOnTopHint
            )
            self.setCursor(Qt.CursorShape.OpenHandCursor)
            self.resize(950, 850)
            self.setMinimumSize(750, 550)
        else:
            # Allow shrinking when embedded
            from PySide6.QtWidgets import QSizePolicy
            self.setMinimumSize(1, 1)
            self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self._build_ui()
        self._analyze_and_display()

        # Apply blur
        try:
            apply_macos_blur(self)
        except:
            pass

        self.raise_()
        self.activateWindow()
        self.show()

    # --------------------------------------------------------
    # Drag window
    # --------------------------------------------------------
    def _drag_start(self, e):
        if e.button() == Qt.LeftButton:
            self._drag_offset = (
                e.globalPosition().toPoint() -
                self.frameGeometry().topLeft()
            )

    def _drag_move(self, e):
        if e.buttons() & Qt.LeftButton:
            self.move(e.globalPosition().toPoint() - self._drag_offset)

    def mousePressEvent(self, event):
        """Enable dragging from anywhere on the panel."""
        if self.embedded:
            return super().mousePressEvent(event)
        if event.button() == Qt.LeftButton:
            self._drag_offset = (
                event.globalPosition().toPoint() -
                self.frameGeometry().topLeft()
            )
            self._dragging = True
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        """Move panel when dragging."""
        if self.embedded:
            return super().mouseMoveEvent(event)
        if getattr(self, '_dragging', False) and event.buttons() & Qt.LeftButton:
            self.move(event.globalPosition().toPoint() - self._drag_offset)
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        """Stop dragging."""
        if self.embedded:
            return super().mouseReleaseEvent(event)
        if event.button() == Qt.LeftButton:
            self._dragging = False
            self.setCursor(Qt.CursorShape.OpenHandCursor)
        super().mouseReleaseEvent(event)

    # --------------------------------------------------------
    # UI BUILD - Matching Patient History Panel
    # --------------------------------------------------------
    def _build_ui(self):
        # Different styles for embedded vs floating
        if self.embedded:
            self.setStyleSheet("""
                QWidget {
                    background-color: white;
                    color: #333;
                }
                QLabel {
                    background: transparent;
                    color: #333;
                }
                QPushButton {
                    background-color: rgba(0,0,0,0.08);
                    color: #333;
                    border-radius: 6px;
                }
                QPushButton:hover {
                    background-color: rgba(0,0,0,0.15);
                }
            """)
        else:
            self.setStyleSheet("""
                QWidget {
                    background-color: rgba(32,32,32,0.25);
                    color: #DCE6FF;
                    border-radius: 12px;
                }
                QLabel {
                    color: #DCE6FF;
                }
                QPushButton {
                    background-color: rgba(255,255,255,0.22);
                    color: white;
                    border-radius: 6px;
                }
                QPushButton:hover {
                    background-color: rgba(255,255,255,0.35);
                }
            """)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(12, 12, 12, 12)
        outer.setSpacing(12)

        # Background wrapper
        self.bg = QWidget()
        if self.embedded:
            self.bg.setStyleSheet("background-color: white; border-radius: 0;")
        else:
            self.bg.setStyleSheet("""
                background-color: rgba(20,20,20,0.18);
                border-radius: 12px;
            """)
        bg_layout = QVBoxLayout(self.bg)
        bg_layout.setContentsMargins(0, 0, 0, 0)
        bg_layout.setSpacing(0)
        outer.addWidget(self.bg)

        # Title bar
        self.title_bar = QWidget()
        self.title_bar.setFixedHeight(46)
        if not self.embedded:
            self.title_bar.setCursor(Qt.CursorShape.OpenHandCursor)
            self.title_bar.setStyleSheet("""
                background-color: rgba(30,30,30,0.35);
                border-top-left-radius: 12px;
                border-top-right-radius: 12px;
            """)
        else:
            self.title_bar.setStyleSheet("""
                background-color: rgba(240,242,245,0.95);
                border-bottom: 1px solid #d0d5da;
            """)

        tb = QHBoxLayout(self.title_bar)
        tb.setContentsMargins(12, 4, 12, 4)

        title = QLabel("⚠️ Risk Overview")
        if self.embedded:
            title.setStyleSheet("font-size: 18px; font-weight: bold; color: #333; background: transparent;")
        else:
            title.setStyleSheet("font-size: 20px; font-weight: bold; color: #F5F5F5;")
        tb.addWidget(title)

        # Summary label in title bar
        self.summary_label = QLabel("")
        if self.embedded:
            self.summary_label.setStyleSheet("font-size: 13px; color: #666; background: transparent;")
        else:
            self.summary_label.setStyleSheet("font-size: 14px; color: #AAA;")
        tb.addWidget(self.summary_label)

        tb.addStretch()

        # Export button
        export_btn = QPushButton("Export to Word")
        export_btn.setFixedHeight(28)
        export_btn.clicked.connect(self._export_to_word)
        export_btn.setStyleSheet("""
            QPushButton {
                background-color: rgba(70,130,180,0.6);
                color: #FFFFFF;
                font-size: 13px;
                padding: 4px 12px;
                border: 1px solid rgba(255,255,255,0.25);
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: rgba(70,130,180,0.8);
            }
        """)
        tb.addWidget(export_btn)

        # Only add close button for floating mode
        if not self.embedded:
            close_btn = QPushButton("✕")
            close_btn.setFixedSize(34, 28)
            close_btn.clicked.connect(self.close)
            close_btn.setStyleSheet("""
                QPushButton {
                    background-color: rgba(255,255,255,0.18);
                    color: #FFFFFF;
                    font-size: 18px;
                    font-weight: bold;
                    border: 1px solid rgba(255,255,255,0.25);
                    border-radius: 6px;
                }
                QPushButton:hover {
                    background-color: rgba(255,255,255,0.32);
                }
            """)
            tb.addWidget(close_btn)

            self.title_bar.mousePressEvent = self._drag_start
            self.title_bar.mouseMoveEvent = self._drag_move

        bg_layout.addWidget(self.title_bar)

        # Scroll area
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        if self.embedded:
            self.scroll.setStyleSheet("""
                QScrollArea {
                    background-color: white;
                    border: none;
                }
                QScrollBar:vertical {
                    background: rgba(0,0,0,0.05);
                    width: 10px;
                    border-radius: 5px;
                }
                QScrollBar::handle:vertical {
                    background: rgba(0,0,0,0.2);
                    border-radius: 5px;
                    min-height: 30px;
                }
            """)
            self.scroll.viewport().setStyleSheet("background-color: white;")
        else:
            self.scroll.setStyleSheet("""
                QScrollArea {
                    background-color: rgba(0,0,0,0);
                }
                QScrollBar:vertical {
                    background: rgba(255,255,255,0.10);
                    width: 12px;
                    margin: 0px;
                    border-radius: 6px;
                }
                QScrollBar::handle:vertical {
                    background: rgba(255,255,255,0.35);
                    border-radius: 6px;
                    min-height: 40px;
                }
            """)
            self.scroll.viewport().setStyleSheet("""
                background-color: rgba(32,32,32,0.18);
                border-radius: 12px;
            """)

        bg_layout.addWidget(self.scroll)

        # Inner content panel
        self.inner = QWidget()
        self.inner.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.inner.setMinimumWidth(1)
        if self.embedded:
            self.inner.setStyleSheet("background-color: white;")
        else:
            self.inner.setStyleSheet("""
                background-color: rgba(32,32,32,0.22);
                border-radius: 12px;
            """)

        self.inner_layout = QVBoxLayout(self.inner)
        self.inner_layout.setAlignment(Qt.AlignTop)
        self.inner_layout.setSpacing(18)
        self.inner_layout.setContentsMargins(12, 12, 12, 30)

        self.scroll.setWidget(self.inner)

        # Resize grip - only for floating mode
        if not self.embedded:
            self.resize_grip = QSizeGrip(self)
            self.resize_grip.setStyleSheet("""
                background-color: rgba(255,255,255,0.35);
                border-radius: 6px;
                width: 16px;
                height: 16px;
            """)
            bg_layout.addWidget(self.resize_grip, alignment=Qt.AlignBottom | Qt.AlignRight)
        else:
            self.resize_grip = None

    # --------------------------------------------------------
    # ANALYZE AND DISPLAY
    # --------------------------------------------------------
    def _analyze_and_display(self):
        self.results = analyze_notes_for_risk(self.notes)  # Store for export
        self.category_sections = {}  # Store references to sections for scroll-to
        self._current_category_filter = None  # Active category filter for timeline chart

        total = self.results["total_notes"]
        with_incidents = self.results["notes_with_incidents"]
        pct = (with_incidents / total * 100) if total > 0 else 0

        # Build summary with severity counts
        severity_counts = self.results.get("severity_counts", {"high": 0, "medium": 0, "low": 0})
        high_count = severity_counts.get("high", 0)

        summary_text = f"• {total} notes analyzed • {with_incidents} with risk ({pct:.1f}%)"
        if high_count > 0:
            summary_text += f" • {high_count} HIGH severity"

        self.summary_label.setText(summary_text)

        # =====================================================
        # CHART 1: Risk Overview (pie chart with interactive legend)
        # =====================================================
        chart_bytes, category_info = create_pie_chart(self.results["categories"], embedded=self.embedded)
        if chart_bytes:
            risk_overview_section = CollapsibleSection("📊 Risk Overview", start_collapsed=True, embedded=self.embedded)
            self.inner_layout.addWidget(risk_overview_section)

            self.pie_chart = PieChartWithLegend(embedded=self.embedded)
            self.pie_chart.set_data(chart_bytes, category_info, self.results["categories"])
            self.pie_chart.incidentSelected.connect(self._scroll_to_note_in_panel)

            pie_chart_scroll = ResizableChartContainer(initial_height=380, min_height=280, max_height=550)
            pie_chart_scroll.viewport().setStyleSheet("background: transparent;")
            pie_chart_scroll.setWidget(self.pie_chart)
            risk_overview_section.add_widget(pie_chart_scroll)

        # =====================================================
        # CHART 2: Risk Incidents Over Time (timeline)
        # =====================================================
        incidents_section = CollapsibleSection("📊 Risk Incidents Over Time", start_collapsed=True, embedded=self.embedded)
        self.inner_layout.addWidget(incidents_section)

        # Controls row (outside the resizable container)
        controls_widget = QWidget()
        controls_widget.setStyleSheet("background: transparent;")
        controls_layout = QHBoxLayout(controls_widget)
        controls_layout.setContentsMargins(4, 4, 4, 8)
        controls_layout.setSpacing(6)

        # Button style based on embedded mode
        if self.embedded:
            zoom_btn_style = """
                QPushButton {
                    background-color: rgba(0,0,0,0.08);
                    color: #333;
                    font-size: 11px;
                    border-radius: 4px;
                    border: 1px solid rgba(0,0,0,0.15);
                }
                QPushButton:hover { background-color: rgba(70,130,180,0.3); }
                QPushButton:checked { background-color: rgba(70,130,180,0.5); color: white; }
            """
            label_style = "font-size: 12px; color: #555; background: transparent;"
        else:
            zoom_btn_style = """
                QPushButton {
                    background-color: rgba(255,255,255,0.15);
                    color: white;
                    font-size: 11px;
                    border-radius: 4px;
                    border: 1px solid rgba(255,255,255,0.2);
                }
                QPushButton:hover { background-color: rgba(70,130,180,0.5); }
                QPushButton:checked { background-color: rgba(70,130,180,0.7); }
            """
            label_style = "font-size: 12px; color: #AAA;"

        self.zoom_buttons = []
        self.zoom_levels = [("3M", 0.25), ("6M", 0.5), ("1Y", 1), ("2Y", 2), ("5Y", 5), ("All", None)]
        for label, years in self.zoom_levels:
            btn = QPushButton(label)
            btn.setFixedSize(40, 24)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setStyleSheet(zoom_btn_style)
            btn.setCheckable(True)
            if label == "All":
                btn.setChecked(True)
            btn.clicked.connect(lambda checked, y=years: self._update_timeline_zoom(y))
            controls_layout.addWidget(btn)
            self.zoom_buttons.append(btn)

        self.current_zoom_index = len(self.zoom_levels) - 1

        controls_layout.addSpacing(20)

        # Visual zoom controls
        minus_btn = QPushButton("−")
        minus_btn.setFixedSize(28, 24)
        minus_btn.setCursor(Qt.PointingHandCursor)
        minus_btn.setStyleSheet(zoom_btn_style)
        minus_btn.clicked.connect(self._zoom_out_visual)
        controls_layout.addWidget(minus_btn)

        self.zoom_indicator = QLabel("100%")
        self.zoom_indicator.setStyleSheet(label_style + " min-width: 40px;")
        self.zoom_indicator.setAlignment(Qt.AlignCenter)
        controls_layout.addWidget(self.zoom_indicator)

        plus_btn = QPushButton("+")
        plus_btn.setFixedSize(28, 24)
        plus_btn.setCursor(Qt.PointingHandCursor)
        plus_btn.setStyleSheet(zoom_btn_style)
        plus_btn.clicked.connect(self._zoom_in_visual)
        controls_layout.addWidget(plus_btn)

        reset_btn = QPushButton("Reset")
        reset_btn.setFixedSize(45, 24)
        reset_btn.setCursor(Qt.PointingHandCursor)
        reset_btn.setStyleSheet(zoom_btn_style)
        reset_btn.clicked.connect(self._reset_visual_zoom)
        controls_layout.addWidget(reset_btn)

        controls_layout.addStretch()
        incidents_section.add_widget(controls_widget)

        # ---- Category filter buttons ----
        self._build_category_filters(incidents_section)

        # Timeline chart
        self.timeline_chart = InteractiveTimelineChart()
        self.timeline_chart.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self.timeline_chart.setStyleSheet("background: transparent;")
        self.timeline_chart.incidentSelected.connect(self._scroll_to_note_in_panel)
        self.timeline_chart.zoomChanged.connect(self._on_visual_zoom_changed)

        timeline_scroll = ResizableChartContainer(initial_height=220, min_height=120, max_height=500)
        timeline_scroll.viewport().setStyleSheet("background: transparent;")
        timeline_scroll.setWidget(self.timeline_chart)
        incidents_section.add_widget(timeline_scroll)

        # Generate initial timeline (all data)
        self._update_timeline_zoom(None)

        # =====================================================
        # CHART 3: Risk Level Timeline Visual
        # =====================================================
        timeline_visual_section = CollapsibleSection("📊 Risk Level Timeline", start_collapsed=True, embedded=self.embedded)
        self.inner_layout.addWidget(timeline_visual_section)

        timeline_visual_bytes, timeline_visual_info = create_risk_timeline_visual(self.results)
        if timeline_visual_bytes:
            self.risk_level_timeline = InteractiveRiskLevelTimeline()
            self.risk_level_timeline.set_chart(timeline_visual_bytes, timeline_visual_info, self.results["categories"])
            self.risk_level_timeline.setAlignment(Qt.AlignLeft)
            self.risk_level_timeline.setStyleSheet("background: transparent; padding: 8px;")
            self.risk_level_timeline.incidentSelected.connect(self._scroll_to_note_in_panel)

            risk_timeline_scroll = ResizableChartContainer(initial_height=180, min_height=100, max_height=400)
            risk_timeline_scroll.viewport().setStyleSheet("background: transparent;")
            risk_timeline_scroll.setWidget(self.risk_level_timeline)
            timeline_visual_section.add_widget(risk_timeline_scroll)
        else:
            no_timeline_label = QLabel("Insufficient data for timeline visualization")
            if self.embedded:
                no_timeline_label.setStyleSheet("color: #666; padding: 16px; background: transparent;")
            else:
                no_timeline_label.setStyleSheet("color: #888; padding: 16px;")
            no_timeline_label.setAlignment(Qt.AlignCenter)
            timeline_visual_section.add_widget(no_timeline_label)

        # Severity summary section
        severity_counts = self.results.get("severity_counts", {"high": 0, "medium": 0, "low": 0})
        if any(severity_counts.values()):
            severity_widget = QWidget()
            severity_widget.setStyleSheet("background: transparent;")
            severity_layout = QHBoxLayout(severity_widget)
            severity_layout.setContentsMargins(0, 8, 0, 16)
            severity_layout.setSpacing(20)

            severity_label = QLabel("Severity:")
            severity_label.setStyleSheet("font-size: 13px; color: #AAA;")
            severity_layout.addWidget(severity_label)

            # High severity
            high_count = severity_counts.get("high", 0)
            if high_count > 0:
                high_pill = QLabel(f"HIGH: {high_count}")
                high_pill.setStyleSheet("""
                    background-color: rgba(220,38,38,0.7);
                    color: white;
                    font-size: 12px;
                    font-weight: bold;
                    padding: 4px 12px;
                    border-radius: 12px;
                """)
                severity_layout.addWidget(high_pill)

            # Medium severity
            med_count = severity_counts.get("medium", 0)
            if med_count > 0:
                med_pill = QLabel(f"MEDIUM: {med_count}")
                med_pill.setStyleSheet("""
                    background-color: rgba(245,158,11,0.7);
                    color: white;
                    font-size: 12px;
                    font-weight: bold;
                    padding: 4px 12px;
                    border-radius: 12px;
                """)
                severity_layout.addWidget(med_pill)

            # Low severity
            low_count = severity_counts.get("low", 0)
            if low_count > 0:
                low_pill = QLabel(f"LOW: {low_count}")
                low_pill.setStyleSheet("""
                    background-color: rgba(34,197,94,0.7);
                    color: white;
                    font-size: 12px;
                    font-weight: bold;
                    padding: 4px 12px;
                    border-radius: 12px;
                """)
                severity_layout.addWidget(low_pill)

            severity_layout.addStretch()
            self.inner_layout.addWidget(severity_widget)

        # Sort by count
        sorted_cats = sorted(
            self.results["categories"].items(),
            key=lambda x: x[1]["count"],
            reverse=True
        )

        # =====================================================
        # RISK DETAILS - Collapsible container for all categories
        # =====================================================
        total_incidents = sum(cat["count"] for cat in self.results["categories"].values())
        risk_details_section = CollapsibleSection(
            f"📋 Risk Details ({total_incidents} incidents)",
            start_collapsed=True,
            embedded=self.embedded
        )
        self.inner_layout.addWidget(risk_details_section)

        for cat_name, cat_data in sorted_cats:
            count = cat_data["count"]
            if count == 0:
                continue

            icon = cat_data["icon"]
            color = cat_data["color"]
            incidents = cat_data["incidents"]
            patterns = cat_data["patterns"]
            subcategories = cat_data.get("subcategories", {})

            # Create collapsible section with severity indicator
            # Count high severity incidents in this category
            high_in_cat = sum(1 for inc in incidents if inc.get("severity") == "high")
            severity_indicator = ""
            if high_in_cat > 0:
                severity_indicator = f" <span style='color: #ef4444;'>[{high_in_cat} HIGH]</span>"

            title = f"{icon}  {cat_name} ({count}){severity_indicator}"
            section = CollapsibleSection(title, start_collapsed=True, embedded=self.embedded)
            risk_details_section.add_widget(section)
            self.category_sections[cat_name] = section  # Store reference for scroll-to

            # Add subcategory breakdown if available
            if subcategories:
                sorted_subcats = sorted(
                    [(name, data) for name, data in subcategories.items() if data["count"] > 0],
                    key=lambda x: x[1]["count"],
                    reverse=True
                )

                if sorted_subcats:
                    # Subcategory pills - use a grid layout to allow wrapping
                    subcat_widget = QWidget()
                    subcat_widget.setStyleSheet("background: transparent;")
                    from PySide6.QtWidgets import QGridLayout
                    subcat_layout = QGridLayout(subcat_widget)
                    subcat_layout.setContentsMargins(8, 4, 8, 12)
                    subcat_layout.setSpacing(8)

                    # Store containers for each subcategory
                    subcat_containers = {}

                    row, col = 0, 0
                    max_cols = 3  # Max 3 pills per row

                    for subcat_name, subcat_data in sorted_subcats[:6]:  # Limit to 6 subcategories
                        subcat_count = subcat_data["count"]
                        subcat_severity = subcat_data.get("severity", "medium")

                        # Severity border color
                        if subcat_severity == "high":
                            border_color = "#ef4444"
                        elif subcat_severity == "medium":
                            border_color = "#f59e0b"
                        else:
                            border_color = "#22c55e"

                        pill = QPushButton(f"{subcat_name}: {subcat_count}")
                        pill.setCursor(Qt.PointingHandCursor)
                        pill.setMinimumHeight(32)
                        if self.embedded:
                            pill.setStyleSheet(f"""
                                QPushButton {{
                                    background-color: rgba(240,242,245,0.95);
                                    color: #333;
                                    font-size: 12px;
                                    padding: 6px 14px;
                                    border-radius: 12px;
                                    border: 2px solid {border_color};
                                    text-align: left;
                                }}
                                QPushButton:hover {{
                                    background-color: rgba(200,220,240,0.95);
                                }}
                                QPushButton:checked {{
                                    background-color: rgba(100,150,200,0.3);
                                }}
                            """)
                        else:
                            pill.setStyleSheet(f"""
                                QPushButton {{
                                    background-color: rgba(60,60,60,0.8);
                                    color: white;
                                    font-size: 12px;
                                    padding: 6px 14px;
                                    border-radius: 12px;
                                    border: 2px solid {border_color};
                                    text-align: left;
                                }}
                                QPushButton:hover {{
                                    background-color: rgba(80,80,80,0.9);
                                }}
                                QPushButton:checked {{
                                    background-color: rgba(70,130,180,0.7);
                                }}
                            """)
                        pill.setCheckable(True)
                        subcat_layout.addWidget(pill, row, col)
                        col += 1
                        if col >= max_cols:
                            col = 0
                            row += 1

                        # Create container for this subcategory's incidents (hidden by default)
                        container = QWidget()
                        container.setStyleSheet("background: transparent;")
                        container_layout = QVBoxLayout(container)
                        container_layout.setContentsMargins(8, 8, 8, 8)
                        container_layout.setSpacing(8)
                        container.hide()
                        subcat_containers[subcat_name] = container

                        # Connect pill click to toggle container
                        def make_toggle(cont, btn, all_containers):
                            def toggle(checked):
                                # Hide all other containers
                                for name, c in all_containers.items():
                                    if c is not cont:
                                        c.hide()
                                # Toggle this one
                                cont.setVisible(checked)
                                # Uncheck other buttons
                                for b in btn.parent().findChildren(QPushButton):
                                    if b is not btn:
                                        b.setChecked(False)
                            return toggle
                        pill.clicked.connect(make_toggle(container, pill, subcat_containers))

                    section.add_widget(subcat_widget)

                    # Add incident containers and populate them
                    for subcat_name, subcat_data in sorted_subcats[:6]:
                        container = subcat_containers[subcat_name]
                        section.add_widget(container)

                        # Deduplicate by date - keep most recent incident per date
                        subcat_incidents = subcat_data["incidents"]
                        seen_dates = {}
                        for inc in sorted(subcat_incidents, key=lambda x: x.get("date") or datetime.min, reverse=True):
                            date_key = inc["date"].strftime("%Y-%m-%d") if inc.get("date") else "Unknown"
                            if date_key not in seen_dates:
                                seen_dates[date_key] = inc
                        deduped_incidents = sorted(seen_dates.values(), key=lambda x: x.get("date") or datetime.min, reverse=True)

                        for inc in deduped_incidents:
                            date_str = inc["date"].strftime("%d %b %Y") if inc.get("date") else "Unknown"
                            full_text = inc["full_text"]

                            # Get patterns for this specific subcategory
                            subcat_patterns = RISK_CATEGORIES.get(cat_name, {}).get("subcategories", {}).get(subcat_name, {}).get("patterns", patterns)

                            # Highlight matched patterns
                            highlighted_text = highlight_matches(full_text, subcat_patterns)

                            # Date collapsible section
                            date_section = CollapsibleSection(date_str, start_collapsed=True, embedded=self.embedded)
                            container.layout().addWidget(date_section)

                            # Full text block with highlighting
                            text_block = QLabel(highlighted_text)
                            text_block.setWordWrap(True)
                            text_block.setTextFormat(Qt.RichText)
                            if self.embedded:
                                text_block.setStyleSheet("""
                                    padding: 10px 12px;
                                    background-color: #f8f9fa;
                                    border-radius: 10px;
                                    color: #333;
                                    font-size: 13px;
                                    line-height: 1.4;
                                """)
                            else:
                                text_block.setStyleSheet("""
                                    padding: 10px 12px;
                                    background-color: rgba(40,40,40,0.72);
                                    border-radius: 10px;
                                    color: #DCE6FF;
                                    font-size: 13px;
                                    line-height: 1.4;
                                """)
                            text_block.mousePressEvent = lambda e, d=inc.get("date"): self._scroll_to_note_in_panel({"date": d})
                            text_block.setCursor(Qt.PointingHandCursor)
                            date_section.add_widget(text_block)

            else:
                # No subcategories - show incidents directly (fallback)
                # Deduplicate by date - keep most recent incident per date
                seen_dates = {}
                for inc in sorted(incidents, key=lambda x: x.get("date") or datetime.min, reverse=True):
                    date_key = inc["date"].strftime("%Y-%m-%d") if inc.get("date") else "Unknown"
                    if date_key not in seen_dates:
                        seen_dates[date_key] = inc
                deduped_incidents = sorted(seen_dates.values(), key=lambda x: x.get("date") or datetime.min, reverse=True)

                for inc in deduped_incidents:
                    date_str = inc["date"].strftime("%d %b %Y") if inc.get("date") else "Unknown"
                    full_text = inc["full_text"]

                    # Highlight matched patterns
                    highlighted_text = highlight_matches(full_text, patterns)

                    # Date collapsible section
                    date_section = CollapsibleSection(date_str, start_collapsed=True, embedded=self.embedded)
                    section.add_widget(date_section)

                    # Full text block with highlighting
                    text_block = QLabel(highlighted_text)
                    text_block.setWordWrap(True)
                    text_block.setTextFormat(Qt.RichText)
                    if self.embedded:
                        text_block.setStyleSheet("""
                            padding: 10px 12px;
                            background-color: #f8f9fa;
                            border-radius: 10px;
                            color: #333;
                            font-size: 13px;
                            line-height: 1.4;
                        """)
                    else:
                        text_block.setStyleSheet("""
                            padding: 10px 12px;
                            background-color: rgba(40,40,40,0.72);
                            border-radius: 10px;
                            color: #DCE6FF;
                            font-size: 13px;
                            line-height: 1.4;
                        """)
                    text_block.mousePressEvent = lambda e, d=inc.get("date"): self._scroll_to_note_in_panel({"date": d})
                    text_block.setCursor(Qt.PointingHandCursor)
                    date_section.add_widget(text_block)

        if all(cat["count"] == 0 for cat in self.results["categories"].values()):
            no_risk = QLabel("No risk incidents detected in the notes.")
            no_risk.setStyleSheet("font-size: 16px; color: #AAA;")
            no_risk.setAlignment(Qt.AlignCenter)
            self.inner_layout.addWidget(no_risk)

        self.inner_layout.addStretch(1)

    # --------------------------------------------------------
    # EXPORT TO WORD
    # --------------------------------------------------------
    def _export_to_word(self):
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        import os

        # Get save path
        default_name = f"Risk_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Risk Report", default_name,
            "Word Document (*.docx)"
        )

        if not file_path:
            return

        # Export
        result = export_risk_to_word(self.results, file_path)
        if result:
            show_styled_message(self, "Export Complete", f"Report saved to:\n{file_path}")
            # Open the file
            import platform
            if platform.system() == 'Windows':
                os.startfile(file_path)
            else:
                os.system(f'open "{file_path}"')
        else:
            show_styled_message(self, "Export Failed", "Could not export report. Check if python-docx is installed.", is_warning=True)

    # --------------------------------------------------------
    # SCROLL TO NOTE IN LEFT PANEL
    # --------------------------------------------------------
    def _scroll_to_note_in_panel(self, incident: dict):
        """Scroll to the specific note in the left notes panel and highlight matched text."""
        if not self.notes_panel:
            return

        date = incident.get("date")
        matched_text = incident.get("matched", "")

        if date:
            try:
                # Jump to date in notes panel
                self.notes_panel.jump_to_date(date)

                # Highlight the matched text if the notes panel supports it
                if hasattr(self.notes_panel, 'highlight_text') and matched_text:
                    self.notes_panel.highlight_text(matched_text)
                elif hasattr(self.notes_panel, 'search_and_highlight') and matched_text:
                    self.notes_panel.search_and_highlight(matched_text)

            except Exception as e:
                print(f"Error scrolling to note: {e}")

    # --------------------------------------------------------
    # EXPORT SINGLE CATEGORY TO WORD
    # --------------------------------------------------------
    def _export_category_to_word(self, category_name: str):
        """Export a single category's incidents to Word."""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        import os

        if category_name not in self.results["categories"]:
            return

        cat_data = self.results["categories"][category_name]
        if cat_data["count"] == 0:
            show_styled_message(self, "No Data", f"No incidents found for {category_name}")
            return

        # Get save path
        safe_name = category_name.replace("/", "-").replace(" ", "_")
        default_name = f"Risk_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(
            self, f"Export {category_name}", default_name,
            "Word Document (*.docx)"
        )

        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading(f'{cat_data["icon"]} {category_name} Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            summary = doc.add_paragraph()
            summary.add_run(f"Total incidents: {cat_data['count']}\n").bold = True

            # Incidents
            doc.add_heading('Incidents', level=1)

            for inc in cat_data["incidents"]:
                date_str = inc["date"].strftime("%d %b %Y") if inc["date"] else "Unknown"
                p = doc.add_paragraph()
                p.add_run(f"{date_str}: ").bold = True

                text = inc.get("full_text", "")[:800]
                if len(inc.get("full_text", "")) > 800:
                    text += "..."
                p.add_run(text)
                p.paragraph_format.space_after = Pt(8)

            doc.save(file_path)

            show_styled_message(self, "Export Complete", f"{category_name} report saved to:\n{file_path}")
            import platform
            if platform.system() == 'Windows':
                os.startfile(file_path)
            else:
                os.system(f'open "{file_path}"')

        except Exception as e:
            show_styled_message(self, "Export Failed", f"Could not export: {e}", is_warning=True)

    # --------------------------------------------------------
    # SCROLL TO CATEGORY
    # --------------------------------------------------------
    def _scroll_to_category(self, category_name: str):
        """Scroll to and expand the specified category section."""
        if category_name not in self.category_sections:
            return

        section = self.category_sections[category_name]

        # Expand the section if collapsed (check container visibility)
        if not section.container.isVisible():
            section._toggle()

        # Scroll to make section visible
        self.scroll.ensureWidgetVisible(section, 50, 50)

        # Brief highlight effect on header
        original_style = section.header_bar.styleSheet()
        section.header_bar.setStyleSheet("background-color: rgba(74,144,217,0.4); border-radius: 6px;")

        # Remove highlight after delay
        from PySide6.QtCore import QTimer
        QTimer.singleShot(1000, lambda: section.header_bar.setStyleSheet(original_style))

    # --------------------------------------------------------
    # VISUAL ZOOM CONTROLS
    # --------------------------------------------------------
    def _on_visual_zoom_changed(self, percentage):
        """Update zoom indicator when visual zoom changes."""
        self.zoom_indicator.setText(f"{percentage}%")

    def _zoom_in_visual(self):
        """Zoom in on the timeline chart."""
        self.timeline_chart.zoom_in()

    def _zoom_out_visual(self):
        """Zoom out on the timeline chart."""
        self.timeline_chart.zoom_out()

    def _reset_visual_zoom(self):
        """Reset visual zoom to 100%."""
        self.timeline_chart.reset_zoom()

    # --------------------------------------------------------
    # CATEGORY FILTER for timeline chart
    # --------------------------------------------------------
    def _build_category_filters(self, parent_section):
        """Build category filter buttons for the timeline chart."""
        # Collect active categories with their colors
        active_cats = [
            (name, info)
            for name, info in sorted(self.results["categories"].items())
            if info["count"] > 0
        ]
        if not active_cats:
            return

        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(4, 0, 4, 4)
        filter_layout.setSpacing(4)

        # Scrollable row of category buttons
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(36)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollBar:horizontal {
                height: 5px; background: rgba(0,0,0,0.05); border-radius: 2px;
            }
            QScrollBar::handle:horizontal {
                background: rgba(0,0,0,0.2); border-radius: 2px; min-width: 20px;
            }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(5)

        filter_label = QLabel("Filter:")
        if self.embedded:
            filter_label.setStyleSheet("font-size: 11px; color: #666; background: transparent;")
        else:
            filter_label.setStyleSheet("font-size: 11px; color: #AAA; background: transparent;")
        label_row.addWidget(filter_label)

        self._category_filter_buttons = []
        for cat_name, cat_info in active_cats:
            color = cat_info["color"]
            icon = cat_info["icon"]
            count = cat_info["count"]
            btn = QPushButton(f"{icon} {cat_name} ({count})")
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 4px 10px;
                    font-size: 11px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background: {color}dd;
                }}
            """)
            btn.setCursor(Qt.PointingHandCursor)
            btn.setProperty("cat_name", cat_name)
            btn.setProperty("cat_color", color)
            btn.clicked.connect(lambda checked, c=cat_name: self._apply_category_filter(c))
            label_row.addWidget(btn)
            self._category_filter_buttons.append(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row (hidden initially)
        self._cat_filter_status = QWidget()
        self._cat_filter_status.setStyleSheet("background: transparent;")
        self._cat_filter_status.setVisible(False)
        status_layout = QHBoxLayout(self._cat_filter_status)
        status_layout.setContentsMargins(0, 0, 0, 0)
        status_layout.setSpacing(6)

        self._cat_filter_label = QLabel("")
        if self.embedded:
            self._cat_filter_label.setStyleSheet(
                "font-size: 12px; color: #374151; font-weight: 500; background: transparent;"
            )
        else:
            self._cat_filter_label.setStyleSheet(
                "font-size: 12px; color: #DCE6FF; font-weight: 500; background: transparent;"
            )
        status_layout.addWidget(self._cat_filter_label)

        remove_btn = QPushButton("✕ Clear")
        remove_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444; color: white; border: none;
                border-radius: 4px; padding: 3px 8px; font-size: 11px; font-weight: 600;
            }
            QPushButton:hover { background: #dc2626; }
        """)
        remove_btn.setCursor(Qt.PointingHandCursor)
        remove_btn.clicked.connect(self._remove_category_filter)
        status_layout.addWidget(remove_btn)
        status_layout.addStretch()
        filter_layout.addWidget(self._cat_filter_status)

        parent_section.add_widget(filter_container)

    def _apply_category_filter(self, category_name: str):
        """Filter the timeline chart to show only a specific category."""
        self._current_category_filter = category_name

        # Find color/icon for this category
        cat_info = self.results["categories"].get(category_name, {})
        icon = cat_info.get("icon", "")
        self._cat_filter_label.setText(f"Showing: {icon} {category_name}")
        self._cat_filter_status.setVisible(True)

        # Re-render chart with current time zoom + category filter
        years = self.zoom_levels[self.current_zoom_index][1]
        self._update_timeline_zoom(years)

    def _remove_category_filter(self):
        """Remove category filter and show all categories."""
        self._current_category_filter = None
        self._cat_filter_status.setVisible(False)

        # Re-render chart with current time zoom, no category filter
        years = self.zoom_levels[self.current_zoom_index][1]
        self._update_timeline_zoom(years)

    # --------------------------------------------------------
    # TIMELINE DATA ZOOM
    # --------------------------------------------------------
    def _update_timeline_zoom(self, years):
        """Update timeline chart to show only the specified number of years (None = all)."""
        # Update button checked states
        year_map = {0.25: 0, 0.5: 1, 1: 2, 2: 3, 5: 4, None: 5}
        btn_index = year_map.get(years, 5)
        self.current_zoom_index = btn_index
        for i, btn in enumerate(self.zoom_buttons):
            btn.setChecked(i == btn_index)

        # Get monthly counts
        monthly_counts = self.results["monthly_counts"]

        if not monthly_counts:
            self.timeline_chart.setText("No timeline data available")
            return

        # Filter data based on years
        if years is None:
            # Show all data
            filtered_counts = dict(monthly_counts)
        else:
            # Calculate cutoff date
            cutoff = datetime.now() - timedelta(days=years * 365)
            cutoff_key = cutoff.strftime("%Y-%m")

            # Filter to only months after cutoff
            filtered_counts = {
                month: counts
                for month, counts in monthly_counts.items()
                if month >= cutoff_key
            }

        # Apply category filter if active
        cat_filter = getattr(self, '_current_category_filter', None)
        if cat_filter:
            cat_filtered = {}
            for month, counts in filtered_counts.items():
                if cat_filter in counts:
                    cat_filtered[month] = {cat_filter: counts[cat_filter]}
            filtered_counts = cat_filtered

        if not filtered_counts:
            if cat_filter and years:
                self.timeline_chart.setText(
                    f"No {cat_filter} data in the last {years} year(s)"
                )
            elif cat_filter:
                self.timeline_chart.setText(f"No {cat_filter} data available")
            else:
                self.timeline_chart.setText(f"No data in the last {years} year(s)")
            return

        # Build categories_data: when filtering, only pass the filtered category
        categories_data = self.results["categories"]
        if cat_filter:
            categories_data = {
                k: v for k, v in categories_data.items()
                if k == cat_filter
            }

        # Compute global y_max from the full (unfiltered) monthly data so
        # filtered charts stay proportional to the overall scale
        all_monthly = self.results["monthly_counts"]
        global_y_max = 0
        for month_cats in all_monthly.values():
            for c in month_cats.values():
                if c > global_y_max:
                    global_y_max = c

        # Use the category's own colour when a filter is active
        color_override = None
        if cat_filter:
            color_override = self.results["categories"].get(cat_filter, {}).get("color")

        # Generate chart with filtered data (returns tuple with timeline_info)
        chart_bytes, timeline_info = create_timeline_chart(
            filtered_counts, categories_data,
            y_max=global_y_max, color_override=color_override
        )
        if chart_bytes:
            self.timeline_chart.set_chart(chart_bytes, timeline_info, self.results["categories"])
        else:
            self.timeline_chart.setText("Unable to generate timeline chart")

    def closeEvent(self, event):
        self.closed.emit()
        super().closeEvent(event)
