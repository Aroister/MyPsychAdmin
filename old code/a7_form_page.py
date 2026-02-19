# ================================================================
#  A7 FORM PAGE — Section 3 Joint Medical Recommendation for Treatment
#  Mental Health Act 1983 - Form A7 Regulation 4(1)(d)(i)
#  Single scrollable form layout
# ================================================================

from __future__ import annotations

from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QGroupBox, QToolButton, QRadioButton,
    QButtonGroup, QComboBox, QSpinBox, QCompleter, QStyleFactory,
    QSlider
)

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# TOOLBAR
# ================================================================

class A7Toolbar(QWidget):
    """Toolbar for the A7 Form Page."""

    export_docx = Signal()
    import_file = Signal()
    clear_form = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(60)
        self.setStyleSheet("""
            A7Toolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 8, 16, 8)
        layout.setSpacing(12)

        # Export DOCX button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #7c3aed;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #6d28d9; }
            QToolButton:pressed { background: #5b21b6; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Import File button
        import_btn = QToolButton()
        import_btn.setText("Import File")
        import_btn.setFixedSize(100, 38)
        import_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #059669; }
            QToolButton:pressed { background: #047857; }
        """)
        import_btn.clicked.connect(self.import_file.emit)
        layout.addWidget(import_btn)

        # Clear Form button
        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #dc2626; }
            QToolButton:pressed { background: #b91c1c; }
        """)
        clear_btn.clicked.connect(self.clear_form.emit)
        layout.addWidget(clear_btn)

        layout.addStretch()


# ================================================================
# MAIN A7 FORM PAGE
# ================================================================

class A7FormPage(QWidget):
    """Page for completing MHA Form A7 - Section 3 Joint Medical Recommendation for Treatment."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean",
        "Asian",
        "Caucasian",
        "Middle Eastern",
        "Mixed Race",
        "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()

        self._setup_ui()
        self._prefill_first_practitioner()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {
            "full_name": details[1] or "",
            "email": details[7] or "",
        }

    def _prefill_first_practitioner(self):
        if self._my_details.get("full_name"):
            self.prac1_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.prac1_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #7c3aed; border-bottom: 1px solid #6d28d9;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Form A7 — Section 3 Joint Medical Recommendation for Treatment")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = A7Toolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        self.toolbar.import_file.connect(self._import_file)
        self.toolbar.clear_form.connect(self._clear_form)
        main_layout.addWidget(self.toolbar)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        # Build all sections
        self._create_section_1_demographics()
        self._create_section_2_patient()
        self._create_section_3_first_practitioner()
        self._create_section_4_second_practitioner()
        self._create_section_5_clinical_reasons_with_controls()
        self._create_section_6_appropriate_treatment()
        self._create_section_7_signatures()

        self.form_layout.addStretch()

        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, number: int, title: str) -> QFrame:
        """Create a styled section frame with number and title."""
        frame = QFrame()
        frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)

        # Section header
        header_layout = QHBoxLayout()
        header_layout.setSpacing(12)

        number_badge = QLabel(str(number))
        number_badge.setFixedSize(32, 32)
        number_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        number_badge.setStyleSheet("""
            background: #7c3aed;
            color: white;
            font-size: 14px;
            font-weight: 700;
            border-radius: 16px;
        """)
        header_layout.addWidget(number_badge)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1f2937;")
        header_layout.addWidget(title_lbl)
        header_layout.addStretch()

        layout.addLayout(header_layout)

        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus { border-color: #7c3aed; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 80) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMaximumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #7c3aed; }
        """)
        return edit

    def _create_form_row(self, label_text: str, widget: QWidget, label_width: int = 140) -> QHBoxLayout:
        row = QHBoxLayout()
        row.setSpacing(12)
        lbl = QLabel(label_text)
        lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        lbl.setFixedWidth(label_width)
        row.addWidget(lbl)
        row.addWidget(widget, 1)
        return row

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        return date_edit

    # ----------------------------------------------------------------
    # SECTION 1: Demographics (Age, Gender, Ethnicity) - Single Line
    # ----------------------------------------------------------------
    def _create_section_1_demographics(self):
        frame = self._create_section_frame(1, "Patient Demographics")
        layout = frame.layout()

        # All on one line: Age spinbox, Gender radios, Ethnicity dropdown
        row = QHBoxLayout()
        row.setSpacing(16)

        # Age
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row.addWidget(age_lbl)

        self.age_spin = QSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setFixedWidth(70)
        self.age_spin.setStyleSheet("""
            QSpinBox {
                padding: 8px;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                font-size: 13px;
            }
        """)
        row.addWidget(self.age_spin)

        row.addSpacing(20)

        # Gender (no label)
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("Male")
        self.gender_female = QRadioButton("Female")
        self.gender_other = QRadioButton("Other")
        self.gender_group.addButton(self.gender_male, 0)
        self.gender_group.addButton(self.gender_female, 1)
        self.gender_group.addButton(self.gender_other, 2)

        row.addWidget(self.gender_male)
        row.addWidget(self.gender_female)
        row.addWidget(self.gender_other)

        row.addSpacing(20)

        # Ethnicity (no label - first item is "Ethnicity")
        self.ethnicity_combo = QComboBox()
        self.ethnicity_combo.addItem("Ethnicity")  # Placeholder as first item
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(200)
        self.ethnicity_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                font-size: 13px;
            }
        """)
        row.addWidget(self.ethnicity_combo)

        row.addStretch()
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 2: Patient Details - Single Line
    # ----------------------------------------------------------------
    def _create_section_2_patient(self):
        frame = self._create_section_frame(2, "Patient Details")
        layout = frame.layout()

        # Both on one line, no labels (placeholders are self-explanatory)
        row = QHBoxLayout()
        row.setSpacing(12)

        self.patient_name = self._create_line_edit("Patient's full name")
        row.addWidget(self.patient_name, 1)

        self.patient_address = self._create_line_edit("Patient's address")
        row.addWidget(self.patient_address, 2)

        layout.addLayout(row)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 3: First Practitioner - Compact Layout
    # ----------------------------------------------------------------
    def _create_section_3_first_practitioner(self):
        frame = self._create_section_frame(3, "First Medical Practitioner")
        layout = frame.layout()

        # Row 1: Name, Address, Email (no labels)
        row1 = QHBoxLayout()
        row1.setSpacing(12)

        self.prac1_name = self._create_line_edit("Full name")
        row1.addWidget(self.prac1_name, 1)

        self.prac1_address = self._create_line_edit("Address")
        row1.addWidget(self.prac1_address, 2)

        self.prac1_email = self._create_line_edit("Email")
        row1.addWidget(self.prac1_email, 1)

        layout.addLayout(row1)

        # Row 2: Exam date + checkboxes
        row2 = QHBoxLayout()
        row2.setSpacing(12)

        exam_lbl = QLabel("Examined:")
        exam_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row2.addWidget(exam_lbl)

        self.prac1_exam_date = self._create_date_edit()
        self.prac1_exam_date.setFixedWidth(140)
        row2.addWidget(self.prac1_exam_date)

        row2.addSpacing(20)

        self.prac1_acquaintance = QCheckBox("Previous acquaintance")
        self.prac1_acquaintance.setStyleSheet("font-size: 13px; color: #374151;")
        row2.addWidget(self.prac1_acquaintance)

        self.prac1_section12 = QCheckBox("Section 12 approved")
        self.prac1_section12.setStyleSheet("font-size: 13px; color: #374151;")
        row2.addWidget(self.prac1_section12)

        row2.addStretch()
        layout.addLayout(row2)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 4: Second Practitioner - Compact Layout
    # ----------------------------------------------------------------
    def _create_section_4_second_practitioner(self):
        frame = self._create_section_frame(4, "Second Medical Practitioner")
        layout = frame.layout()

        # Row 1: Name, Address, Email (no labels)
        row1 = QHBoxLayout()
        row1.setSpacing(12)

        self.prac2_name = self._create_line_edit("Full name")
        row1.addWidget(self.prac2_name, 1)

        self.prac2_address = self._create_line_edit("Address")
        row1.addWidget(self.prac2_address, 2)

        self.prac2_email = self._create_line_edit("Email")
        row1.addWidget(self.prac2_email, 1)

        layout.addLayout(row1)

        # Row 2: Exam date + checkboxes
        row2 = QHBoxLayout()
        row2.setSpacing(12)

        exam_lbl = QLabel("Examined:")
        exam_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row2.addWidget(exam_lbl)

        self.prac2_exam_date = self._create_date_edit()
        self.prac2_exam_date.setFixedWidth(140)
        row2.addWidget(self.prac2_exam_date)

        row2.addSpacing(20)

        self.prac2_acquaintance = QCheckBox("Previous acquaintance")
        self.prac2_acquaintance.setStyleSheet("font-size: 13px; color: #374151;")
        row2.addWidget(self.prac2_acquaintance)

        self.prac2_section12 = QCheckBox("Section 12 approved")
        self.prac2_section12.setStyleSheet("font-size: 13px; color: #374151;")
        row2.addWidget(self.prac2_section12)

        row2.addStretch()
        layout.addLayout(row2)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 5: Clinical Reasons with Controls Panel (Split Layout)
    # ----------------------------------------------------------------
    def _create_section_5_clinical_reasons_with_controls(self):
        frame = self._create_section_frame(5, "Clinical Reasons")
        layout = frame.layout()

        # Horizontal split: Clinical text (left) | Controls (right)
        split_layout = QHBoxLayout()
        split_layout.setSpacing(20)

        # === LEFT: Clinical Reasons Text Area ===
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(8)

        info = QLabel("Click options on the right to auto-generate text:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 11px; color: #6b7280; padding: 6px; background: #f3e8ff; border-radius: 4px;")
        left_layout.addWidget(info)

        self.clinical_reasons = QTextEdit()
        self.clinical_reasons.setPlaceholderText("Clinical reasons will be generated here...")
        self.clinical_reasons.setMinimumHeight(400)
        self.clinical_reasons.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 12px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #7c3aed; }
        """)
        left_layout.addWidget(self.clinical_reasons)

        split_layout.addWidget(left_container, 3)

        # === RIGHT: Controls Panel (scrollable) ===
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        right_scroll.setFixedWidth(400)

        right_container = QWidget()
        right_container.setFixedWidth(380)
        right_container.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 8, 0)
        right_layout.setSpacing(12)

        # --- Mental Disorder (ICD-10) - 2 boxes ---
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(8)

        md_header = QLabel("Mental Disorder")
        md_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        self.dx_boxes = []
        for i in range(2):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText("Primary diagnosis..." if i == 0 else "Secondary (optional)...")

            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(ICD10_DICT.items(), key=lambda x: x[0].lower()):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})

            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(12)
            combo.setStyleSheet("""
                QComboBox { padding: 6px; font-size: 12px; border: 1px solid #d1d5db; border-radius: 4px; background: white; }
                QComboBox QAbstractItemView { min-width: 300px; }
            """)
            combo.currentIndexChanged.connect(self._update_clinical_text)
            md_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        right_layout.addWidget(md_frame)

        # --- Legal Criteria ---
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(6)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature with sub-options
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._update_clinical_text)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._update_clinical_text)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._update_clinical_text)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree with slider
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(120)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 12px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 12px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_clinical_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity section
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health with sub-options
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._update_clinical_text)
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._update_clinical_text)
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 12px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._update_clinical_text)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety with sub-options
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # === SELF SECTION ===
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        # Self - Historical
        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._update_clinical_text)
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Self placement in risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._update_clinical_text)
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._update_clinical_text)
        self_opt_layout.addWidget(self.self_hist_harm)

        # Self - Current
        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._update_clinical_text)
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Self placement in risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._update_clinical_text)
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._update_clinical_text)
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # === OTHERS SECTION ===
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        # Others - Historical
        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence to others")
        self.others_hist_violence.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_hist_verbal)

        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_sexual.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_sexual.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_hist_sexual)

        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_stalking.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_stalking.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_hist_stalking)

        self.others_hist_arson = QCheckBox("Arson")
        self.others_hist_arson.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_arson.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_hist_arson)

        # Others - Current
        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence to others")
        self.others_curr_violence.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_sexual.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_sexual.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_curr_sexual)

        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_stalking.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_stalking.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_curr_stalking)

        self.others_curr_arson = QCheckBox("Arson")
        self.others_curr_arson.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_arson.toggled.connect(self._update_clinical_text)
        others_opt_layout.addWidget(self.others_curr_arson)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        right_layout.addWidget(lc_frame)

        # --- Informal Not Indicated ---
        inf_frame = QFrame()
        inf_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; }")
        inf_layout = QVBoxLayout(inf_frame)
        inf_layout.setContentsMargins(12, 10, 12, 10)
        inf_layout.setSpacing(4)

        inf_header = QLabel("Informal Not Indicated")
        inf_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #991b1b;")
        inf_layout.addWidget(inf_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed")
        self.tried_failed_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.tried_failed_cb.toggled.connect(self._update_clinical_text)
        inf_layout.addWidget(self.tried_failed_cb)

        self.insight_cb = QCheckBox("Lack of Insight")
        self.insight_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.insight_cb.toggled.connect(self._update_clinical_text)
        inf_layout.addWidget(self.insight_cb)

        self.compliance_cb = QCheckBox("Compliance Issues")
        self.compliance_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.compliance_cb.toggled.connect(self._update_clinical_text)
        inf_layout.addWidget(self.compliance_cb)

        self.supervision_cb = QCheckBox("Needs MHA Supervision")
        self.supervision_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.supervision_cb.toggled.connect(self._update_clinical_text)
        inf_layout.addWidget(self.supervision_cb)

        right_layout.addWidget(inf_frame)
        right_layout.addStretch()

        right_scroll.setWidget(right_container)
        split_layout.addWidget(right_scroll)

        layout.addLayout(split_layout)
        self.form_layout.addWidget(frame)

    # --- Control toggle handlers ---
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_clinical_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_clinical_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_clinical_text()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_clinical_text()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_clinical_text()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_clinical_text()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_clinical_text()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            self.self_hist_neglect.setChecked(False)
            self.self_hist_risky.setChecked(False)
            self.self_hist_harm.setChecked(False)
            self.self_curr_neglect.setChecked(False)
            self.self_curr_risky.setChecked(False)
            self.self_curr_harm.setChecked(False)
        self._update_clinical_text()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            self.others_hist_violence.setChecked(False)
            self.others_hist_verbal.setChecked(False)
            self.others_hist_sexual.setChecked(False)
            self.others_hist_stalking.setChecked(False)
            self.others_hist_arson.setChecked(False)
            self.others_curr_violence.setChecked(False)
            self.others_curr_verbal.setChecked(False)
            self.others_curr_sexual.setChecked(False)
            self.others_curr_stalking.setChecked(False)
            self.others_curr_arson.setChecked(False)
        self._update_clinical_text()

    def _update_clinical_text(self):
        """Update clinical reasons text based on control selections."""
        self.clinical_reasons.setPlainText(self._generate_clinical_text())

    # ----------------------------------------------------------------
    # SECTION 6: Appropriate Treatment (new for A7/Section 3)
    # ----------------------------------------------------------------
    def _create_section_6_appropriate_treatment(self):
        frame = self._create_section_frame(6, "Appropriate Treatment")
        layout = frame.layout()

        info = QLabel("Specify the hospital(s) where appropriate treatment is available:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 12px; color: #6b7280;")
        layout.addWidget(info)

        self.hospital_treatment = self._create_line_edit("Hospital name(s)")
        layout.addWidget(self.hospital_treatment)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 7: Signatures
    # ----------------------------------------------------------------
    def _create_section_7_signatures(self):
        frame = self._create_section_frame(7, "Signatures")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(20)

        # First practitioner
        sig1_lbl = QLabel("First Practitioner Date:")
        sig1_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        row.addWidget(sig1_lbl)

        self.sig1_date = self._create_date_edit()
        self.sig1_date.setFixedWidth(140)
        row.addWidget(self.sig1_date)

        row.addSpacing(40)

        # Second practitioner
        sig2_lbl = QLabel("Second Practitioner Date:")
        sig2_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        row.addWidget(sig2_lbl)

        self.sig2_date = self._create_date_edit()
        self.sig2_date.setFixedWidth(140)
        row.addWidget(self.sig2_date)

        row.addStretch()
        layout.addLayout(row)

        info = QLabel("NOTE: AT LEAST ONE OF THE PRACTITIONERS SIGNING THIS FORM MUST BE APPROVED UNDER SECTION 12 OF THE ACT.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 11px; color: #dc2626; padding: 6px; background: #fef2f2; border-radius: 4px; font-weight: 500;")
        layout.addWidget(info)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _go_back(self):
        self.go_back.emit()

    def _clear_form(self):
        reply = QMessageBox.question(
            self, "Clear Form", "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.ethnicity_combo.setCurrentIndex(0)
            self.patient_name.clear()
            self.patient_address.clear()
            self.prac1_name.clear()
            self.prac1_address.clear()
            self.prac1_email.clear()
            self.prac1_exam_date.setDate(QDate.currentDate())
            self.prac1_acquaintance.setChecked(False)
            self.prac1_section12.setChecked(False)
            self.prac2_name.clear()
            self.prac2_address.clear()
            self.prac2_email.clear()
            self.prac2_exam_date.setDate(QDate.currentDate())
            self.prac2_acquaintance.setChecked(False)
            self.prac2_section12.setChecked(False)
            for combo in self.dx_boxes:
                combo.setCurrentIndex(0)
            # Legal criteria
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.degree_details.clear()
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            # Informal
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.clinical_reasons.clear()
            self.hospital_treatment.clear()
            self.sig1_date.setDate(QDate.currentDate())
            self.sig2_date.setDate(QDate.currentDate())

    def _import_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Import File", "",
            "All Files (*);;Word Documents (*.docx);;Text Files (*.txt)"
        )
        if file_path:
            QMessageBox.information(self, "Import", f"File selected: {file_path}\n\nData extraction coming soon.")

    def _generate_clinical_text(self) -> str:
        """Generate clinical reasons text from form selections."""
        p = self._get_pronouns()

        # === PARAGRAPH 1: Demographics, Diagnosis, Nature/Degree ===
        para1_parts = []

        # Get patient name for opening
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        # Build opening sentence
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")

        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            eth_simple = ethnicity.replace(" British", "").replace(" Other", "").replace("Mixed ", "")
            opening_parts.append(eth_simple)

        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        # Build diagnosis part
        diagnoses = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    diagnoses.append(f"{dx} ({icd})" if icd else dx)

        if diagnoses:
            if opening_parts:
                demo_str = " ".join(opening_parts)
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {joined} which are mental disorders as defined by the Mental Health Act.")
            else:
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} suffers from {joined} which are mental disorders as defined by the Mental Health Act.")

            # Nature/Degree sentence
            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree to warrant detention for treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature to warrant detention for treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree to warrant detention for treatment.")

            # Nature details
            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    nature_str = " and ".join(nature_types)
                    para1_parts.append(f"The nature of the illness is {nature_str}.")

            # Degree details
            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                # Truncate at "disorder" if present
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                # Simplify schizophrenia variants to just "schizophrenia"
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        elif opening_parts:
            demo_str = " ".join(opening_parts)
            para1_parts.append(f"{name_display} is a {demo_str}.")

        # === PARAGRAPH 2: Necessity (Health + Safety) ===
        para2_parts = []

        # Build necessity summary
        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("safety of others")

        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")

        # Health details
        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_reasons = []
            if self.poor_compliance_cb.isChecked():
                mh_reasons.append("non compliance")
            if self.limited_insight_cb.isChecked():
                mh_reasons.append("limited insight")

            if mh_reasons:
                reasons_str = "/".join(mh_reasons)
                para2_parts.append(f"Regarding health we would be concerned about {p['pos_l']} mental health deteriorating due to {reasons_str}.")
            else:
                para2_parts.append(f"Regarding health we would be concerned about {p['pos_l']} mental health deteriorating.")

        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"We are also concerned about {p['pos_l']} physical health: {details}.")
            else:
                para2_parts.append(f"We are also concerned about {p['pos_l']} physical health.")

        # Self safety details
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            # Get reflexive pronoun (himself/herself/themselves)
            if self.gender_male.isChecked():
                reflexive = "himself"
            elif self.gender_female.isChecked():
                reflexive = "herself"
            else:
                reflexive = "themselves"

            # Check each risk type for historical and current
            risk_types = [
                ("self neglect", self.self_hist_neglect.isChecked(), self.self_curr_neglect.isChecked()),
                (f"placing of {reflexive} in risky situations", self.self_hist_risky.isChecked(), self.self_curr_risky.isChecked()),
                ("self harm", self.self_hist_harm.isChecked(), self.self_curr_harm.isChecked()),
            ]

            both_items = []  # historical and current
            hist_only = []   # historical only
            curr_only = []   # current only

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            self_text = f"With respect to {p['pos_l']} own safety we are concerned about"
            parts = []
            if both_items:
                if len(both_items) == 1:
                    parts.append(f"historical and current {both_items[0]}")
                else:
                    parts.append(f"historical and current {', '.join(both_items[:-1])}, and of {both_items[-1]}")
            if hist_only:
                if len(hist_only) == 1:
                    parts.append(f"historical {hist_only[0]}")
                else:
                    parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
            if curr_only:
                if len(curr_only) == 1:
                    parts.append(f"current {curr_only[0]}")
                else:
                    parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")

            if parts:
                self_text += " " + ", and ".join(parts) + "."
            else:
                self_text += f" {p['pos_l']} safety."

            para2_parts.append(self_text)

        # Others safety details
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            # Check each risk type for historical and current
            risk_types = [
                ("violence to others", self.others_hist_violence.isChecked(), self.others_curr_violence.isChecked()),
                ("verbal aggression", self.others_hist_verbal.isChecked(), self.others_curr_verbal.isChecked()),
                ("sexual violence", self.others_hist_sexual.isChecked(), self.others_curr_sexual.isChecked()),
                ("stalking", self.others_hist_stalking.isChecked(), self.others_curr_stalking.isChecked()),
                ("arson", self.others_hist_arson.isChecked(), self.others_curr_arson.isChecked()),
            ]

            both_items = []  # historical and current
            hist_only = []   # historical only
            curr_only = []   # current only

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            others_text = "With respect to risk to others we are concerned about the risk of"
            parts = []
            if both_items:
                parts.append(f"historical and current {', '.join(both_items)}")
            if hist_only:
                parts.append(f"historical {', '.join(hist_only)}")
            if curr_only:
                parts.append(f"current {', '.join(curr_only)}")

            if parts:
                others_text += " " + " and of ".join(parts) + "."
            else:
                others_text += f" {p['pos_l']} potential to cause harm."

            para2_parts.append(others_text)

        # === PARAGRAPH 3: Informal not indicated ===
        para3_parts = []

        if self.tried_failed_cb.isChecked():
            para3_parts.append("Previous attempts at informal admissions have not been successful and we would likewise be concerned about this recurring in this instance hence we do not believe informal admission currently would be appropriate.")

        if self.insight_cb.isChecked():
            para3_parts.append(f"{p['pos']} lack of insight is a significant concern and should {p['subj_l']} be discharged from section, we believe this would significantly impair {p['pos_l']} compliance if informal.")

        if self.compliance_cb.isChecked():
            if para3_parts:
                para3_parts.append(f"Compliance with treatment has also been a significant issue and we do not believe {p['subj_l']} would comply if informal.")
            else:
                para3_parts.append(f"Compliance with treatment has been a significant issue and we do not believe {p['subj_l']} would comply if informal.")

        if self.supervision_cb.isChecked():
            name = patient_name if patient_name else "the patient"
            para3_parts.append(f"We believe {name} needs careful community monitoring under the supervision afforded by the mental health act and we do not believe such supervision would be complied with should {p['subj_l']} remain in the community informally.")

        # === BUILD FINAL TEXT ===
        paragraphs = []
        if para1_parts:
            paragraphs.append(" ".join(para1_parts))
        if para2_parts:
            paragraphs.append(" ".join(para2_parts))
        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        return "\n\n".join(paragraphs)

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form A7",
            f"Form_A7_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_A7_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form A7 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    run = para.add_run(new_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            paragraphs = doc.paragraphs

            # Data goes into WHITESPACE paragraphs below instruction paragraphs
            # Para 3: Patient data (below para 2 instruction)
            # Para 6: Prac1 data (below para 5 instruction)
            # Para 8: Prac1 date (below para 7 instruction)
            # Para 13: Prac2 data (below para 12 instruction)
            # Para 15: Prac2 date (below para 14 instruction)
            # Para 30: Clinical text (below para 29 instruction)
            # Para 37: Hospital data (below para 36 instruction)

            # Patient - goes in para 3
            patient_text = self.patient_name.text()
            if self.patient_address.text():
                patient_text += ", " + self.patient_address.text()
            if patient_text.strip():
                set_para_text(paragraphs[3], patient_text)
                highlight_yellow(paragraphs[3])

            # First practitioner - goes in para 6
            prac1_text = self.prac1_name.text()
            if self.prac1_address.text():
                prac1_text += ", " + self.prac1_address.text()
            if self.prac1_email.text():
                prac1_text += ", Email: " + self.prac1_email.text()
            if prac1_text.strip():
                set_para_text(paragraphs[6], prac1_text)
                highlight_yellow(paragraphs[6])

            # Prac1 exam date - goes in para 8
            prac1_date = self.prac1_exam_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[8], prac1_date)
            highlight_yellow(paragraphs[8])

            # Prac1 checkboxes (para 9, 10)
            if not self.prac1_acquaintance.isChecked():
                strikethrough_para(paragraphs[9])
            if not self.prac1_section12.isChecked():
                strikethrough_para(paragraphs[10])

            # Second practitioner - goes in para 13
            prac2_text = self.prac2_name.text()
            if self.prac2_address.text():
                prac2_text += ", " + self.prac2_address.text()
            if self.prac2_email.text():
                prac2_text += ", Email: " + self.prac2_email.text()
            if prac2_text.strip():
                set_para_text(paragraphs[13], prac2_text)
                highlight_yellow(paragraphs[13])

            # Prac2 exam date - goes in para 15
            prac2_date = self.prac2_exam_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[15], prac2_date)
            highlight_yellow(paragraphs[15])

            # Prac2 checkboxes (para 16, 17)
            if not self.prac2_acquaintance.isChecked():
                strikethrough_para(paragraphs[16])
            if not self.prac2_section12.isChecked():
                strikethrough_para(paragraphs[17])

            # Detention reasons - strikethrough unselected options
            if not self.health_cb.isChecked():
                strikethrough_para(paragraphs[23])
            if not (self.safety_cb.isChecked() and self.self_harm_cb.isChecked()):
                strikethrough_para(paragraphs[24])
            if not (self.safety_cb.isChecked() and self.others_cb.isChecked()):
                strikethrough_para(paragraphs[25])

            # Clinical reasons - goes in para 30
            clinical_text = self._generate_clinical_text()
            if clinical_text:
                set_para_text(paragraphs[30], clinical_text)
                highlight_yellow(paragraphs[30])

            # Hospital - goes in para 37 (auto-fill from prac1 address if not specified)
            hospital_text = self.hospital_treatment.text().strip()
            if not hospital_text and self.prac1_address.text().strip():
                hospital_text = self.prac1_address.text().strip()
            if hospital_text:
                set_para_text(paragraphs[37], hospital_text)
                highlight_yellow(paragraphs[37])

            # Signatures
            sig1_date = self.sig1_date.date().toString("dd MMMM yyyy")
            sig2_date = self.sig2_date.date().toString("dd MMMM yyyy")
            for run in paragraphs[40].runs:
                run.text = ""
            paragraphs[40].add_run(f"Signed                                                                    Date {sig1_date}")
            for run in paragraphs[41].runs:
                run.text = ""
            paragraphs[41].add_run(f"Signed                                                                    Date {sig2_date}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form A7 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        dx_list = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta:
                dx_list.append(meta)

        gender = "neutral"
        if self.gender_male.isChecked():
            gender = "male"
        elif self.gender_female.isChecked():
            gender = "female"

        return {
            "age": self.age_spin.value(),
            "gender": gender,
            "ethnicity": self.ethnicity_combo.currentText(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "prac1_name": self.prac1_name.text(),
            "prac1_address": self.prac1_address.text(),
            "prac1_email": self.prac1_email.text(),
            "prac1_exam_date": self.prac1_exam_date.date().toString("yyyy-MM-dd"),
            "prac1_acquaintance": self.prac1_acquaintance.isChecked(),
            "prac1_section12": self.prac1_section12.isChecked(),
            "prac2_name": self.prac2_name.text(),
            "prac2_address": self.prac2_address.text(),
            "prac2_email": self.prac2_email.text(),
            "prac2_exam_date": self.prac2_exam_date.date().toString("yyyy-MM-dd"),
            "prac2_acquaintance": self.prac2_acquaintance.isChecked(),
            "prac2_section12": self.prac2_section12.isChecked(),
            "diagnoses": dx_list,
            # Legal criteria
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_hist_neglect": self.self_hist_neglect.isChecked(),
            "self_hist_risky": self.self_hist_risky.isChecked(),
            "self_hist_harm": self.self_hist_harm.isChecked(),
            "self_curr_neglect": self.self_curr_neglect.isChecked(),
            "self_curr_risky": self.self_curr_risky.isChecked(),
            "self_curr_harm": self.self_curr_harm.isChecked(),
            "others": self.others_cb.isChecked(),
            "others_hist_violence": self.others_hist_violence.isChecked(),
            "others_hist_verbal": self.others_hist_verbal.isChecked(),
            "others_hist_sexual": self.others_hist_sexual.isChecked(),
            "others_hist_stalking": self.others_hist_stalking.isChecked(),
            "others_hist_arson": self.others_hist_arson.isChecked(),
            "others_curr_violence": self.others_curr_violence.isChecked(),
            "others_curr_verbal": self.others_curr_verbal.isChecked(),
            "others_curr_sexual": self.others_curr_sexual.isChecked(),
            "others_curr_stalking": self.others_curr_stalking.isChecked(),
            "others_curr_arson": self.others_curr_arson.isChecked(),
            # Informal
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "clinical_reasons": self.clinical_reasons.toPlainText(),
            "hospital_treatment": self.hospital_treatment.text(),
            "sig1_date": self.sig1_date.date().toString("yyyy-MM-dd"),
            "sig2_date": self.sig2_date.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        if not state:
            return

        self.age_spin.setValue(state.get("age", 0))
        g = state.get("gender", "neutral")
        if g == "male":
            self.gender_male.setChecked(True)
        elif g == "female":
            self.gender_female.setChecked(True)
        else:
            self.gender_other.setChecked(True)

        idx = self.ethnicity_combo.findText(state.get("ethnicity", "Not specified"))
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)

        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.prac1_name.setText(state.get("prac1_name", ""))
        self.prac1_address.setText(state.get("prac1_address", ""))
        self.prac1_email.setText(state.get("prac1_email", ""))
        if state.get("prac1_exam_date"):
            self.prac1_exam_date.setDate(QDate.fromString(state["prac1_exam_date"], "yyyy-MM-dd"))
        self.prac1_acquaintance.setChecked(state.get("prac1_acquaintance", False))
        self.prac1_section12.setChecked(state.get("prac1_section12", False))

        self.prac2_name.setText(state.get("prac2_name", ""))
        self.prac2_address.setText(state.get("prac2_address", ""))
        self.prac2_email.setText(state.get("prac2_email", ""))
        if state.get("prac2_exam_date"):
            self.prac2_exam_date.setDate(QDate.fromString(state["prac2_exam_date"], "yyyy-MM-dd"))
        self.prac2_acquaintance.setChecked(state.get("prac2_acquaintance", False))
        self.prac2_section12.setChecked(state.get("prac2_section12", False))

        # Restore diagnoses
        for combo, meta in zip(self.dx_boxes, state.get("diagnoses", [])):
            if meta:
                index = combo.findText(meta.get("diagnosis", ""))
                if index >= 0:
                    combo.setCurrentIndex(index)

        # Legal criteria
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_hist_neglect.setChecked(state.get("self_hist_neglect", False))
        self.self_hist_risky.setChecked(state.get("self_hist_risky", False))
        self.self_hist_harm.setChecked(state.get("self_hist_harm", False))
        self.self_curr_neglect.setChecked(state.get("self_curr_neglect", False))
        self.self_curr_risky.setChecked(state.get("self_curr_risky", False))
        self.self_curr_harm.setChecked(state.get("self_curr_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.others_hist_violence.setChecked(state.get("others_hist_violence", False))
        self.others_hist_verbal.setChecked(state.get("others_hist_verbal", False))
        self.others_hist_sexual.setChecked(state.get("others_hist_sexual", False))
        self.others_hist_stalking.setChecked(state.get("others_hist_stalking", False))
        self.others_hist_arson.setChecked(state.get("others_hist_arson", False))
        self.others_curr_violence.setChecked(state.get("others_curr_violence", False))
        self.others_curr_verbal.setChecked(state.get("others_curr_verbal", False))
        self.others_curr_sexual.setChecked(state.get("others_curr_sexual", False))
        self.others_curr_stalking.setChecked(state.get("others_curr_stalking", False))
        self.others_curr_arson.setChecked(state.get("others_curr_arson", False))
        # Informal
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        self.clinical_reasons.setPlainText(state.get("clinical_reasons", ""))
        self.hospital_treatment.setText(state.get("hospital_treatment", ""))

        if state.get("sig1_date"):
            self.sig1_date.setDate(QDate.fromString(state["sig1_date"], "yyyy-MM-dd"))
        if state.get("sig2_date"):
            self.sig2_date.setDate(QDate.fromString(state["sig2_date"], "yyyy-MM-dd"))
