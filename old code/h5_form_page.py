# ================================================================
#  H5 FORM PAGE — Renewal of Authority for Detention
#  Mental Health Act 1983 - Form H5 Regulation 13(1), (2) and (3)
#  Section 20 — Renewal of authority for detention
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QSpinBox, QRadioButton, QButtonGroup, QComboBox, QCompleter,
    QStyleFactory, QSlider
)

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


class H5Toolbar(QWidget):
    """Toolbar for H5 Form Page."""

    export_docx = Signal()
    clear_form = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(60)
        self.setStyleSheet("""
            H5Toolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 8, 16, 8)
        layout.setSpacing(12)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #dc2626;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #b91c1c; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #6b7280;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #4b5563; }
        """)
        clear_btn.clicked.connect(self.clear_form.emit)
        layout.addWidget(clear_btn)

        layout.addStretch()


class H5FormPage(QWidget):
    """Page for completing MHA Form H5 - Renewal of Authority for Detention."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race", "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {"full_name": details[1] or "", "email": details[7] or ""}

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #dc2626; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form H5 — Renewal of Authority for Detention")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = H5Toolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        self.toolbar.clear_form.connect(self._clear_form)
        main_layout.addWidget(self.toolbar)

        # Single scrollable form area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(20)

        self._create_section_1_patient_rc()
        self._create_section_2_clinical_reasons()
        self._create_section_3_informal()
        self._create_section_4_professional()
        self._create_section_5_signature()

        self.form_layout.addStretch()
        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, number: int, title: str) -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 12px; }")
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 16, 24, 16)
        layout.setSpacing(12)

        header_layout = QHBoxLayout()
        header_layout.setSpacing(12)

        number_badge = QLabel(str(number))
        number_badge.setFixedSize(28, 28)
        number_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        number_badge.setStyleSheet("background: #dc2626; color: white; font-size: 13px; font-weight: 700; border-radius: 14px;")
        header_layout.addWidget(number_badge)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937;")
        header_layout.addWidget(title_lbl)
        header_layout.addStretch()

        layout.addLayout(header_layout)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 8px 10px; font-size: 12px; }
            QLineEdit:focus { border-color: #dc2626; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 8px 10px; font-size: 12px; }")
        return date_edit

    # ----------------------------------------------------------------
    # SECTION 1: Patient & RC Details - Compact with Demographics
    # ----------------------------------------------------------------
    def _create_section_1_patient_rc(self):
        frame = self._create_section_frame(1, "Patient & Clinician Details")
        layout = frame.layout()

        # Hospital row
        hosp_row = QHBoxLayout()
        hosp_row.setSpacing(12)
        hosp_lbl = QLabel("To managers of:")
        hosp_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        hosp_row.addWidget(hosp_lbl)
        self.hospital = self._create_line_edit("Hospital name and address")
        hosp_row.addWidget(self.hospital, 1)
        layout.addLayout(hosp_row)

        # Patient row with demographics
        pt_row = QHBoxLayout()
        pt_row.setSpacing(12)
        self.patient_name = self._create_line_edit("Patient full name")
        pt_row.addWidget(self.patient_name, 1)

        # Age
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        pt_row.addWidget(age_lbl)
        self.age_spin = QSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setFixedWidth(55)
        self.age_spin.setStyleSheet("QSpinBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 12px; }")
        pt_row.addWidget(self.age_spin)

        # Gender
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("O")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("font-size: 11px;")
            self.gender_group.addButton(rb)
            pt_row.addWidget(rb)

        # Ethnicity
        self.ethnicity_combo = QComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(130)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 11px; }")
        pt_row.addWidget(self.ethnicity_combo)
        layout.addLayout(pt_row)

        # Dates row
        dates_row = QHBoxLayout()
        dates_row.setSpacing(12)
        exam_lbl = QLabel("Examined:")
        exam_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        dates_row.addWidget(exam_lbl)
        self.exam_date = self._create_date_edit()
        self.exam_date.setFixedWidth(120)
        dates_row.addWidget(self.exam_date)

        exp_lbl = QLabel("Expires:")
        exp_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        dates_row.addWidget(exp_lbl)
        self.expiry_date = self._create_date_edit()
        self.expiry_date.setFixedWidth(120)
        dates_row.addWidget(self.expiry_date)
        dates_row.addStretch()
        layout.addLayout(dates_row)

        # Consulted row
        consult_row = QHBoxLayout()
        consult_row.setSpacing(12)
        consult_lbl = QLabel("Consulted:")
        consult_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        consult_row.addWidget(consult_lbl)
        self.consulted_name = self._create_line_edit("Full name")
        self.consulted_name.textChanged.connect(self._update_professional_section)
        consult_row.addWidget(self.consulted_name, 1)

        self.consulted_profession = QComboBox()
        self.consulted_profession.addItem("Select profession...")
        self.consulted_profession.addItems([
            "Registered Mental Health Nurse",
            "Registered Learning Disabilities Nurse",
            "Occupational Therapist",
            "Social Worker",
            "Psychologist"
        ])
        self.consulted_profession.setStyleSheet("QComboBox { padding: 8px 10px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 12px; background: white; }")
        self.consulted_profession.currentIndexChanged.connect(self._update_professional_section)
        consult_row.addWidget(self.consulted_profession, 1)
        layout.addLayout(consult_row)

        # RC row
        rc_row = QHBoxLayout()
        rc_row.setSpacing(12)
        rc_lbl = QLabel("RC:")
        rc_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        rc_row.addWidget(rc_lbl)
        self.rc_name = self._create_line_edit("RC full name")
        rc_row.addWidget(self.rc_name, 1)
        self.rc_profession = self._create_line_edit("Profession")
        self.rc_profession.setText("Consultant Psychiatrist")
        rc_row.addWidget(self.rc_profession, 1)
        layout.addLayout(rc_row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 2: Clinical Reasons with Controls Panel (Split Layout)
    # ----------------------------------------------------------------
    def _create_section_2_clinical_reasons(self):
        frame = self._create_section_frame(2, "Reasons for Renewal")
        layout = frame.layout()

        split_layout = QHBoxLayout()
        split_layout.setSpacing(16)

        # === LEFT: Clinical Reasons Text Area ===
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(6)

        # Header row with info and expand/collapse buttons
        header_row = QHBoxLayout()
        info = QLabel("Click options on the right to auto-generate text:")
        info.setStyleSheet("font-size: 10px; color: #6b7280; padding: 4px; background: #fef2f2; border-radius: 4px;")
        header_row.addWidget(info)
        header_row.addStretch()

        # Expand/collapse buttons
        self.expand_btn = QPushButton("+")
        self.expand_btn.setFixedSize(24, 24)
        self.expand_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.expand_btn.setStyleSheet("""
            QPushButton { background: #dc2626; color: white; border: none; border-radius: 4px; font-size: 14px; font-weight: bold; }
            QPushButton:hover { background: #b91c1c; }
        """)
        self.expand_btn.clicked.connect(self._expand_reasons)
        header_row.addWidget(self.expand_btn)

        self.collapse_btn = QPushButton("-")
        self.collapse_btn.setFixedSize(24, 24)
        self.collapse_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.collapse_btn.setStyleSheet("""
            QPushButton { background: #6b7280; color: white; border: none; border-radius: 4px; font-size: 14px; font-weight: bold; }
            QPushButton:hover { background: #4b5563; }
        """)
        self.collapse_btn.clicked.connect(self._collapse_reasons)
        header_row.addWidget(self.collapse_btn)

        left_layout.addLayout(header_row)

        self.reasons = QTextEdit()
        self.reasons.setPlaceholderText("Reasons for renewal will be generated here...")
        self._reasons_height = 300  # Track current height
        self.reasons.setFixedHeight(self._reasons_height)
        self.reasons.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px; font-size: 12px; }
            QTextEdit:focus { border-color: #dc2626; }
        """)
        left_layout.addWidget(self.reasons)

        split_layout.addWidget(left_container, 3)

        # === RIGHT: Controls Panel ===
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        right_scroll.setFixedWidth(400)

        right_container = QWidget()
        right_container.setFixedWidth(380)
        right_container.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 8, 0)
        right_layout.setSpacing(10)

        # --- Mental Disorder (ICD-10) ---
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(10, 8, 10, 8)
        md_layout.setSpacing(6)

        md_header = QLabel("Mental Disorder")
        md_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        self.dx_boxes = []
        for i in range(2):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText("Primary diagnosis..." if i == 0 else "Secondary (optional)...")
            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(ICD10_DICT.items(), key=lambda x: x[0].lower()):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})
            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(12)
            combo.setStyleSheet("QComboBox { padding: 5px; font-size: 11px; border: 1px solid #d1d5db; border-radius: 4px; background: white; }")
            combo.currentIndexChanged.connect(self._update_reasons_text)
            md_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        right_layout.addWidget(md_frame)

        # --- Legal Criteria ---
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(10, 8, 10, 8)
        lc_layout.setSpacing(4)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(14, 0, 0, 0)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.chronic_cb = QCheckBox("Chronic and enduring")
        for cb in [self.relapsing_cb, self.treatment_resistant_cb, self.chronic_cb]:
            cb.setStyleSheet("font-size: 11px; color: #6b7280;")
            cb.toggled.connect(self._update_reasons_text)
            nature_opt_layout.addWidget(cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(14, 0, 0, 0)
        degree_opt_layout.setSpacing(3)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setFixedWidth(100)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)
        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 11px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 11px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_reasons_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151; margin-top: 3px;")
        lc_layout.addWidget(nec_lbl)

        # Health with sub-options
        self.nec_health = QCheckBox("Health")
        self.nec_health.setStyleSheet("font-size: 11px; color: #374151;")
        self.nec_health.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.nec_health)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(14, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(14, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._update_reasons_text)
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._update_reasons_text)
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 11px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._update_reasons_text)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety with sub-options
        self.nec_safety = QCheckBox("Safety")
        self.nec_safety.setStyleSheet("font-size: 11px; color: #374151;")
        self.nec_safety.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.nec_safety)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(14, 2, 0, 2)
        safety_opt_layout.setSpacing(2)

        # Self section
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(14, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_hist_harm)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others section
        self.nec_others = QCheckBox("Others")
        self.nec_others.setStyleSheet("font-size: 11px; font-weight: 600; color: #6b7280;")
        self.nec_others.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.nec_others)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(14, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence")
        self.others_hist_violence.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_verbal)

        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_sexual.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_sexual.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_sexual)

        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_stalking.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_stalking.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_stalking)

        self.others_hist_arson = QCheckBox("Arson")
        self.others_hist_arson.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_arson.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_arson)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence")
        self.others_curr_violence.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_sexual.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_sexual.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_sexual)

        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_stalking.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_stalking.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_stalking)

        self.others_curr_arson = QCheckBox("Arson")
        self.others_curr_arson.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_arson.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_arson)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        right_layout.addWidget(lc_frame)
        right_layout.addStretch()

        right_scroll.setWidget(right_container)
        split_layout.addWidget(right_scroll)

        layout.addLayout(split_layout)
        self.form_layout.addWidget(frame)

    def _expand_reasons(self):
        self._reasons_height = min(self._reasons_height + 100, 600)
        self.reasons.setFixedHeight(self._reasons_height)

    def _collapse_reasons(self):
        self._reasons_height = max(self._reasons_height - 100, 200)
        self.reasons.setFixedHeight(self._reasons_height)

    # ----------------------------------------------------------------
    # SECTION 3: Treatment Cannot Be Provided Unless Detained
    # ----------------------------------------------------------------
    def _create_section_3_informal(self):
        frame = self._create_section_frame(3, "Treatment Cannot Be Provided Unless Detained")
        layout = frame.layout()

        split_layout = QHBoxLayout()
        split_layout.setSpacing(16)

        # === LEFT: Text Area ===
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(6)

        info = QLabel("Click options on the right to auto-generate text:")
        info.setStyleSheet("font-size: 10px; color: #6b7280; padding: 4px; background: #fef2f2; border-radius: 4px;")
        left_layout.addWidget(info)

        self.informal_reasons = QTextEdit()
        self.informal_reasons.setPlaceholderText("Reasons why treatment cannot be provided unless detained...")
        self.informal_reasons.setMinimumHeight(150)
        self.informal_reasons.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px; font-size: 12px; }
            QTextEdit:focus { border-color: #dc2626; }
        """)
        left_layout.addWidget(self.informal_reasons)

        split_layout.addWidget(left_container, 3)

        # === RIGHT: Red Controls ===
        inf_frame = QFrame()
        inf_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; }")
        inf_frame.setFixedWidth(280)
        inf_layout = QVBoxLayout(inf_frame)
        inf_layout.setContentsMargins(12, 10, 12, 10)
        inf_layout.setSpacing(6)

        inf_header = QLabel("Why Informal Not Appropriate")
        inf_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #991b1b;")
        inf_layout.addWidget(inf_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed informal")
        self.insight_cb = QCheckBox("Lack of Insight")
        self.compliance_cb = QCheckBox("Compliance Issues")
        self.supervision_cb = QCheckBox("Needs MHA Supervision")
        for cb in [self.tried_failed_cb, self.insight_cb, self.compliance_cb, self.supervision_cb]:
            cb.setStyleSheet("font-size: 11px; color: #374151;")
            cb.toggled.connect(self._update_informal_text)
            inf_layout.addWidget(cb)

        inf_layout.addStretch()
        split_layout.addWidget(inf_frame)

        layout.addLayout(split_layout)
        self.form_layout.addWidget(frame)

    def _update_informal_text(self):
        self.informal_reasons.setPlainText(self._generate_informal_text())

    def _generate_informal_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        parts = []
        if self.tried_failed_cb.isChecked():
            parts.append("Previous attempts at informal treatment have not been successful and I would be concerned about this recurring.")
        if self.insight_cb.isChecked():
            parts.append(f"{p['pos']} lack of insight is a significant concern and should {p['subj_l']} be discharged, I believe this would significantly impair {p['pos_l']} compliance.")
        if self.compliance_cb.isChecked():
            parts.append(f"Compliance with treatment has been a significant issue and I do not believe {p['subj_l']} would comply if informal.")
        if self.supervision_cb.isChecked():
            parts.append(f"I believe {name_display.lower() if name_display != 'The patient' else 'the patient'} needs the supervision afforded by the Mental Health Act.")

        if parts:
            return "Such treatment cannot be provided unless the patient continues to be detained under the Act. " + " ".join(parts)
        return ""

    # --- Control toggle handlers ---
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_reasons_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_reasons_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_reasons_text()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_reasons_text()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_reasons_text()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_reasons_text()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.nec_others.setChecked(False)
        self._update_reasons_text()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            self.self_hist_neglect.setChecked(False)
            self.self_hist_risky.setChecked(False)
            self.self_hist_harm.setChecked(False)
            self.self_curr_neglect.setChecked(False)
            self.self_curr_risky.setChecked(False)
            self.self_curr_harm.setChecked(False)
        self._update_reasons_text()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            self.others_hist_violence.setChecked(False)
            self.others_hist_verbal.setChecked(False)
            self.others_hist_sexual.setChecked(False)
            self.others_hist_stalking.setChecked(False)
            self.others_hist_arson.setChecked(False)
            self.others_curr_violence.setChecked(False)
            self.others_curr_verbal.setChecked(False)
            self.others_curr_sexual.setChecked(False)
            self.others_curr_stalking.setChecked(False)
            self.others_curr_arson.setChecked(False)
        self._update_reasons_text()

    def _update_reasons_text(self):
        self.reasons.setPlainText(self._generate_reasons_text())

    def _update_professional_section(self):
        """Auto-fill professional agreement section when consulted details are filled."""
        name = self.consulted_name.text().strip()
        profession_idx = self.consulted_profession.currentIndex()
        profession = self.consulted_profession.currentText() if profession_idx > 0 else ""

        if name and profession:
            self.prof_name.setText(name)
            self.prof_profession.setText(profession)

    def _generate_reasons_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        paragraphs = []

        # Para 1: Demographics + Diagnosis
        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", "").replace("Mixed ", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    diagnoses.append(f"{dx} ({icd})" if icd else dx)

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for the patient to receive medical treatment in hospital.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for the patient to receive medical treatment in hospital.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for the patient to receive medical treatment in hospital.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                # Truncate at "disorder" if present
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                # Simplify schizophrenia variants to just "schizophrenia"
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        # Para 2: Necessity - improved flowing prose
        para2_parts = []
        necessity_items = []
        if self.nec_health.isChecked():
            necessity_items.append(f"{p['pos_l']} own health")
        if self.nec_safety.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append("own safety")
        if self.nec_safety.isChecked() and self.nec_others.isChecked():
            necessity_items.append("the protection of others")

        if necessity_items:
            joined = ", ".join(necessity_items[:-1]) + f" and {necessity_items[-1]}" if len(necessity_items) > 1 else necessity_items[0]
            para2_parts.append(f"The Mental Health Act is necessary for {joined}.")

        # Health details - flowing prose
        if self.nec_health.isChecked():
            mh_items = []
            if self.mental_health_cb.isChecked():
                if self.poor_compliance_cb.isChecked():
                    mh_items.append("poor compliance with treatment")
                if self.limited_insight_cb.isChecked():
                    mh_items.append(f"limited insight into {p['pos_l']} illness")
            if self.physical_health_cb.isChecked():
                ph_details = self.physical_health_details.text().strip()
                if ph_details:
                    mh_items.append(ph_details)
            if mh_items:
                para2_parts.append(f"Regarding {p['pos_l']} health, I would be concerned about {' and '.join(mh_items)} resulting in a deterioration in mental state if informal.")

        # Safety details - Self - flowing prose
        if self.nec_safety.isChecked() and self.self_harm_cb.isChecked():
            self_hist = []
            self_curr = []
            # Get reflexive pronoun (himself/herself/themselves)
            if self.gender_male.isChecked():
                reflexive = "himself"
            elif self.gender_female.isChecked():
                reflexive = "herself"
            else:
                reflexive = "themselves"
            if self.self_hist_neglect.isChecked(): self_hist.append("self neglect")
            if self.self_hist_risky.isChecked(): self_hist.append(f"placing of {reflexive} in risky situations")
            if self.self_hist_harm.isChecked(): self_hist.append("self harm")
            if self.self_curr_neglect.isChecked(): self_curr.append("self neglect")
            if self.self_curr_risky.isChecked(): self_curr.append(f"placing of {reflexive} in risky situations")
            if self.self_curr_harm.isChecked(): self_curr.append("self harm")

            if self_hist or self_curr:
                safety_text = f"With respect to {p['pos_l']} own safety"
                if self_hist:
                    # Join with "of" before each item and ", and of" before the last
                    if len(self_hist) == 1:
                        hist_joined = self_hist[0]
                    else:
                        hist_joined = ", ".join(self_hist[:-1]) + ", and of " + self_hist[-1]
                    safety_text += f", historically, when unwell and non-compliant there has been a risk of {hist_joined}"
                if self_curr:
                    if self_hist:
                        safety_text += f". This risk continues currently despite treatment"
                    else:
                        # Join with "of" before each item and ", and of" before the last
                        if len(self_curr) == 1:
                            curr_joined = self_curr[0]
                        else:
                            curr_joined = ", ".join(self_curr[:-1]) + ", and of " + self_curr[-1]
                        safety_text += f", currently there is a risk of {curr_joined}"
                safety_text += "."
                para2_parts.append(safety_text)

        # Safety details - Others - flowing prose
        if self.nec_safety.isChecked() and self.nec_others.isChecked():
            others_hist = []
            others_curr = []
            if self.others_hist_violence.isChecked(): others_hist.append("violence")
            if self.others_hist_verbal.isChecked(): others_hist.append("verbal aggression")
            if self.others_hist_sexual.isChecked(): others_hist.append("sexual violence")
            if self.others_hist_stalking.isChecked(): others_hist.append("stalking")
            if self.others_hist_arson.isChecked(): others_hist.append("arson")
            if self.others_curr_violence.isChecked(): others_curr.append("violence")
            if self.others_curr_verbal.isChecked(): others_curr.append("verbal aggression")
            if self.others_curr_sexual.isChecked(): others_curr.append("sexual violence")
            if self.others_curr_stalking.isChecked(): others_curr.append("stalking")
            if self.others_curr_arson.isChecked(): others_curr.append("arson")

            if others_hist or others_curr:
                # Combine items that appear in both historical and current
                both = [item for item in others_hist if item in others_curr]
                hist_only = [item for item in others_hist if item not in both]
                curr_only = [item for item in others_curr if item not in both]

                others_text = "Regarding risk to others, if not under the Mental Health Act and not compliant, I would be concerned about the risk of "
                all_items = []
                if both:
                    all_items.append(f"{', '.join(both)}, which is both historical and current")
                if hist_only:
                    all_items.append(f"{', '.join(hist_only)} (historical)")
                if curr_only:
                    all_items.append(f"{', '.join(curr_only)} (current)")
                # Join with ", and " before the last item if multiple
                if len(all_items) > 1:
                    others_text += ", and ".join([", ".join(all_items[:-1]), all_items[-1]]) + "."
                else:
                    others_text += all_items[0] + "."
                para2_parts.append(others_text)

        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        return "\n\n".join(paragraphs)

    # ----------------------------------------------------------------
    # SECTION 4: Professional Agreement - Compact
    # ----------------------------------------------------------------
    def _create_section_4_professional(self):
        frame = self._create_section_frame(4, "Professional Agreement")
        layout = frame.layout()

        info = QLabel("I agree with the responsible clinician that this patient meets the criteria for renewal of detention.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 11px; color: #374151; padding: 6px; background: #f0fdf4; border-radius: 4px;")
        layout.addWidget(info)

        row = QHBoxLayout()
        row.setSpacing(12)
        self.prof_name = self._create_line_edit("Full name")
        row.addWidget(self.prof_name, 1)
        self.prof_profession = self._create_line_edit("Profession")
        row.addWidget(self.prof_profession, 1)
        sig_lbl = QLabel("Date:")
        sig_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        row.addWidget(sig_lbl)
        self.prof_sig_date = self._create_date_edit()
        self.prof_sig_date.setFixedWidth(120)
        row.addWidget(self.prof_sig_date)
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 5: Signature - Compact
    # ----------------------------------------------------------------
    def _create_section_5_signature(self):
        frame = self._create_section_frame(5, "RC Signature")
        layout = frame.layout()

        # Furnishing report options
        furnish_lbl = QLabel("I am furnishing this report by:")
        furnish_lbl.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        layout.addWidget(furnish_lbl)

        self.furnish_group = QButtonGroup(self)
        self.furnish_internal = QRadioButton("Today consigning it to the hospital managers' internal mail system")
        self.furnish_electronic = QRadioButton("Today sending it to the hospital managers by means of electronic communication")
        self.furnish_other = QRadioButton("Sending or delivering it without using the hospital managers' internal mail system")

        for rb in [self.furnish_internal, self.furnish_electronic, self.furnish_other]:
            rb.setStyleSheet("font-size: 11px; color: #374151; margin-left: 10px;")
            self.furnish_group.addButton(rb)
            layout.addWidget(rb)

        self.furnish_internal.setChecked(True)  # Default selection

        # Signature date row
        row = QHBoxLayout()
        row.setSpacing(20)
        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151; margin-top: 8px;")
        row.addWidget(sig_lbl)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(120)
        row.addWidget(self.rc_sig_date)
        row.addStretch()
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.hospital.clear()
            self.patient_name.clear()
            self.exam_date.setDate(QDate.currentDate())
            self.expiry_date.setDate(QDate.currentDate())
            self.consulted_name.clear()
            self.consulted_profession.setCurrentIndex(0)
            self.rc_name.clear()
            self.rc_profession.setText("Consultant Psychiatrist")
            for combo in self.dx_boxes:
                combo.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.degree_details.clear()
            self.nec_health.setChecked(False)
            self.nec_safety.setChecked(False)
            self.physical_health_details.clear()
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.reasons.clear()
            self.informal_reasons.clear()
            self.prof_name.clear()
            self.prof_profession.clear()
            self.prof_sig_date.setDate(QDate.currentDate())
            self.rc_sig_date.setDate(QDate.currentDate())
            self.furnish_internal.setChecked(True)

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form H5",
            f"Form_H5_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_H5_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form H5 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Hospital (para 5)
            if self.hospital.text().strip():
                set_para_text(paragraphs[5], self.hospital.text())
                highlight_yellow(paragraphs[5])

            # Patient name (para 7)
            if self.patient_name.text().strip():
                set_para_text(paragraphs[7], self.patient_name.text())
                highlight_yellow(paragraphs[7])

            # Exam date (para 9)
            set_para_text(paragraphs[9], self.exam_date.date().toString("dd MMMM yyyy"))
            highlight_yellow(paragraphs[9])

            # Expiry date (para 11)
            set_para_text(paragraphs[11], self.expiry_date.date().toString("dd MMMM yyyy"))
            highlight_yellow(paragraphs[11])

            # Consulted (para 13)
            consult_text = self.consulted_name.text()
            if self.consulted_profession.currentIndex() > 0:
                consult_text += ", " + self.consulted_profession.currentText()
            if consult_text.strip():
                set_para_text(paragraphs[13], consult_text)
                highlight_yellow(paragraphs[13])

            # Necessity (strikethrough unselected)
            # Para 19: "for the patient's own health"
            if not self.nec_health.isChecked():
                strikethrough_para(paragraphs[19])
            # Para 20: "for the patient's own safety" - need Safety AND Self checked
            if not (self.nec_safety.isChecked() and self.self_harm_cb.isChecked()):
                strikethrough_para(paragraphs[20])
            # Para 21: "for the protection of other persons" - need Safety AND Others checked
            if not (self.nec_safety.isChecked() and self.nec_others.isChecked()):
                strikethrough_para(paragraphs[21])

            # Reasons (para 25-27)
            reasons_text = self._generate_reasons_text()
            if reasons_text.strip():
                set_para_text(paragraphs[25], reasons_text)
                highlight_yellow(paragraphs[25])

            # Informal reasons (para 30-32) - use the generated text from section 3
            informal_text = self.informal_reasons.toPlainText().strip()
            if informal_text:
                set_para_text(paragraphs[30], informal_text)
                highlight_yellow(paragraphs[30])

            # PART 1 Signature block (para 36-37)
            # Para 36: "Signed [signature]    PRINT NAME [name]"
            # Para 37: "Profession [profession]    Date [date]"
            rc_name = self.rc_name.text().strip()
            rc_profession = self.rc_profession.text().strip()
            rc_date = self.rc_sig_date.date().toString("dd MMMM yyyy")

            part1_sig = f"Signed                                              PRINT NAME {rc_name}"
            set_para_text(paragraphs[36], part1_sig)
            highlight_yellow(paragraphs[36])

            part1_prof = f"Profession {rc_profession}                                   Date {rc_date}"
            set_para_text(paragraphs[37], part1_prof)
            highlight_yellow(paragraphs[37])

            # PART 2 Signature block (para 41-42)
            prof_name = self.prof_name.text().strip()
            prof_profession = self.prof_profession.text().strip()
            prof_date = self.prof_sig_date.date().toString("dd MMMM yyyy")

            part2_sig = f"Signed                                              PRINT NAME {prof_name}"
            set_para_text(paragraphs[41], part2_sig)
            highlight_yellow(paragraphs[41])

            part2_prof = f"Profession {prof_profession}                                   Date {prof_date}"
            set_para_text(paragraphs[42], part2_prof)
            highlight_yellow(paragraphs[42])

            # PART 3 - Furnishing report options (para 46-48)
            # Strikethrough the options that are NOT selected
            if not self.furnish_internal.isChecked():
                strikethrough_para(paragraphs[46])
            if not self.furnish_electronic.isChecked():
                strikethrough_para(paragraphs[47])
            if not self.furnish_other.isChecked():
                strikethrough_para(paragraphs[48])

            # PART 3 Signature block (para 49-50)
            part3_sig = f"Signed"
            set_para_text(paragraphs[49], part3_sig)

            part3_name = f"PRINT NAME {rc_name}                                   Date {rc_date}"
            set_para_text(paragraphs[50], part3_name)
            highlight_yellow(paragraphs[50])

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form H5 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        gender = "neutral"
        if self.gender_male.isChecked():
            gender = "male"
        elif self.gender_female.isChecked():
            gender = "female"

        dx_list = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta:
                dx_list.append(meta)

        return {
            "age": self.age_spin.value(),
            "gender": gender,
            "ethnicity": self.ethnicity_combo.currentText(),
            "hospital": self.hospital.text(),
            "patient_name": self.patient_name.text(),
            "exam_date": self.exam_date.date().toString("yyyy-MM-dd"),
            "expiry_date": self.expiry_date.date().toString("yyyy-MM-dd"),
            "consulted_name": self.consulted_name.text(),
            "consulted_profession": self.consulted_profession.currentText(),
            "rc_name": self.rc_name.text(),
            "rc_profession": self.rc_profession.text(),
            "diagnoses": dx_list,
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "nec_health": self.nec_health.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "nec_safety": self.nec_safety.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_hist_neglect": self.self_hist_neglect.isChecked(),
            "self_hist_risky": self.self_hist_risky.isChecked(),
            "self_hist_harm": self.self_hist_harm.isChecked(),
            "self_curr_neglect": self.self_curr_neglect.isChecked(),
            "self_curr_risky": self.self_curr_risky.isChecked(),
            "self_curr_harm": self.self_curr_harm.isChecked(),
            "nec_others": self.nec_others.isChecked(),
            "others_hist_violence": self.others_hist_violence.isChecked(),
            "others_hist_verbal": self.others_hist_verbal.isChecked(),
            "others_hist_sexual": self.others_hist_sexual.isChecked(),
            "others_hist_stalking": self.others_hist_stalking.isChecked(),
            "others_hist_arson": self.others_hist_arson.isChecked(),
            "others_curr_violence": self.others_curr_violence.isChecked(),
            "others_curr_verbal": self.others_curr_verbal.isChecked(),
            "others_curr_sexual": self.others_curr_sexual.isChecked(),
            "others_curr_stalking": self.others_curr_stalking.isChecked(),
            "others_curr_arson": self.others_curr_arson.isChecked(),
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "reasons": self.reasons.toPlainText(),
            "informal_reasons": self.informal_reasons.toPlainText(),
            "prof_name": self.prof_name.text(),
            "prof_profession": self.prof_profession.text(),
            "prof_sig_date": self.prof_sig_date.date().toString("yyyy-MM-dd"),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "furnish_method": "internal" if self.furnish_internal.isChecked() else ("electronic" if self.furnish_electronic.isChecked() else "other"),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.age_spin.setValue(state.get("age", 0))
        g = state.get("gender", "neutral")
        if g == "male":
            self.gender_male.setChecked(True)
        elif g == "female":
            self.gender_female.setChecked(True)
        else:
            self.gender_other.setChecked(True)
        idx = self.ethnicity_combo.findText(state.get("ethnicity", "Not specified"))
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.hospital.setText(state.get("hospital", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        if state.get("exam_date"):
            self.exam_date.setDate(QDate.fromString(state["exam_date"], "yyyy-MM-dd"))
        if state.get("expiry_date"):
            self.expiry_date.setDate(QDate.fromString(state["expiry_date"], "yyyy-MM-dd"))
        self.consulted_name.setText(state.get("consulted_name", ""))
        prof_idx = self.consulted_profession.findText(state.get("consulted_profession", ""))
        if prof_idx >= 0:
            self.consulted_profession.setCurrentIndex(prof_idx)
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_profession.setText(state.get("rc_profession", ""))
        for combo, meta in zip(self.dx_boxes, state.get("diagnoses", [])):
            if meta:
                index = combo.findText(meta.get("diagnosis", ""))
                if index >= 0:
                    combo.setCurrentIndex(index)
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.nec_health.setChecked(state.get("nec_health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.nec_safety.setChecked(state.get("nec_safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_hist_neglect.setChecked(state.get("self_hist_neglect", False))
        self.self_hist_risky.setChecked(state.get("self_hist_risky", False))
        self.self_hist_harm.setChecked(state.get("self_hist_harm", False))
        self.self_curr_neglect.setChecked(state.get("self_curr_neglect", False))
        self.self_curr_risky.setChecked(state.get("self_curr_risky", False))
        self.self_curr_harm.setChecked(state.get("self_curr_harm", False))
        self.nec_others.setChecked(state.get("nec_others", False))
        self.others_hist_violence.setChecked(state.get("others_hist_violence", False))
        self.others_hist_verbal.setChecked(state.get("others_hist_verbal", False))
        self.others_hist_sexual.setChecked(state.get("others_hist_sexual", False))
        self.others_hist_stalking.setChecked(state.get("others_hist_stalking", False))
        self.others_hist_arson.setChecked(state.get("others_hist_arson", False))
        self.others_curr_violence.setChecked(state.get("others_curr_violence", False))
        self.others_curr_verbal.setChecked(state.get("others_curr_verbal", False))
        self.others_curr_sexual.setChecked(state.get("others_curr_sexual", False))
        self.others_curr_stalking.setChecked(state.get("others_curr_stalking", False))
        self.others_curr_arson.setChecked(state.get("others_curr_arson", False))
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        self.reasons.setPlainText(state.get("reasons", ""))
        self.informal_reasons.setPlainText(state.get("informal_reasons", ""))
        self.prof_name.setText(state.get("prof_name", ""))
        self.prof_profession.setText(state.get("prof_profession", ""))
        if state.get("prof_sig_date"):
            self.prof_sig_date.setDate(QDate.fromString(state["prof_sig_date"], "yyyy-MM-dd"))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        # Restore furnish method selection
        furnish = state.get("furnish_method", "internal")
        if furnish == "internal":
            self.furnish_internal.setChecked(True)
        elif furnish == "electronic":
            self.furnish_electronic.setChecked(True)
        else:
            self.furnish_other.setChecked(True)
