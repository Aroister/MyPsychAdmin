# ================================================================
#  T2 FORM PAGE — Certificate of Consent to Treatment
#  Mental Health Act 1983 - Form T2 Regulation 27(2)
#  Section 58(3)(a) — Certificate of consent to treatment
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QPushButton, QFileDialog, QMessageBox, QToolButton,
    QComboBox, QRadioButton, QButtonGroup, QSplitter, QGroupBox
)


class T2FormPage(QWidget):
    """Page for completing MHA Form T2 - Certificate of Consent to Treatment."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.clinician_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.clinician_email.setText(self._my_details["email"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #059669; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form T2 — Certificate of Consent to Treatment")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        toolbar = QWidget()
        toolbar.setFixedHeight(60)
        toolbar.setStyleSheet("background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12);")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 8, 16, 8)
        tb_layout.setSpacing(12)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #059669;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #047857; }
        """)
        export_btn.clicked.connect(self._export_docx)
        tb_layout.addWidget(export_btn)

        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #dc2626; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        tb_layout.addWidget(clear_btn)
        tb_layout.addStretch()

        main_layout.addWidget(toolbar)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        self._build_form()

        self.form_layout.addStretch()
        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, title: str) -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 12px; }")
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("font-size: 16px; font-weight: 700; color: #059669;")
        layout.addWidget(title_lbl)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus { border-color: #059669; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 100) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMinimumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #059669; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return date_edit

    def _add_medication_entry(self, med_type: str):
        """Add a medication entry row (regular or prn)."""
        from CANONICAL_MEDS import MEDICATIONS

        entry_widget = QFrame()
        entry_widget.setStyleSheet("""
            QFrame {
                background: #f9fafb;
                border-radius: 4px;
                border: 1px solid #e5e7eb;
            }
        """)
        entry_layout = QHBoxLayout(entry_widget)
        entry_layout.setContentsMargins(8, 6, 8, 6)
        entry_layout.setSpacing(8)

        # Medication dropdown
        med_combo = QComboBox()
        med_combo.setEditable(True)
        med_combo.addItem("")
        med_combo.addItems(sorted(MEDICATIONS.keys()))
        med_combo.setMinimumWidth(140)
        med_combo.setStyleSheet("""
            QComboBox {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 11px;
            }
        """)
        entry_layout.addWidget(med_combo, 1)

        # BNF radio buttons
        bnf_group = QButtonGroup(entry_widget)
        bnf_radio = QRadioButton("BNF")
        bnf_radio.setChecked(True)
        bnf_radio.setStyleSheet("font-size: 11px;")
        above_bnf_radio = QRadioButton("Above BNF")
        above_bnf_radio.setStyleSheet("font-size: 11px;")
        bnf_group.addButton(bnf_radio)
        bnf_group.addButton(above_bnf_radio)
        entry_layout.addWidget(bnf_radio)
        entry_layout.addWidget(above_bnf_radio)

        # Remove button
        remove_btn = QPushButton("×")
        remove_btn.setFixedSize(20, 20)
        remove_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 10px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover { background: #dc2626; }
        """)
        entry_layout.addWidget(remove_btn)

        # Store entry data
        entry_data = {
            "widget": entry_widget,
            "name": med_combo,
            "bnf_radio": bnf_radio,
            "above_bnf_radio": above_bnf_radio
        }

        if med_type == "regular":
            self._regular_meds.append(entry_data)
            self.regular_meds_container.addWidget(entry_widget)
        else:
            self._prn_meds.append(entry_data)
            self.prn_meds_container.addWidget(entry_widget)

        def remove_entry():
            if med_type == "regular" and len(self._regular_meds) > 1:
                self._regular_meds.remove(entry_data)
                entry_widget.deleteLater()
                self._update_treatment_text()
            elif med_type == "prn" and len(self._prn_meds) > 1:
                self._prn_meds.remove(entry_data)
                entry_widget.deleteLater()
                self._update_treatment_text()

        med_combo.currentTextChanged.connect(self._update_treatment_text)
        bnf_radio.toggled.connect(self._update_treatment_text)
        above_bnf_radio.toggled.connect(self._update_treatment_text)
        remove_btn.clicked.connect(remove_entry)

    def _update_treatment_text(self):
        """Generate treatment description based on medication entries."""
        regular_meds = []
        prn_meds = []
        above_bnf_meds = []

        # Collect regular medications
        for entry in self._regular_meds:
            name = entry["name"].currentText().strip()
            if name:
                regular_meds.append(name)
                if entry["above_bnf_radio"].isChecked():
                    above_bnf_meds.append(name)

        # Collect PRN medications
        for entry in self._prn_meds:
            name = entry["name"].currentText().strip()
            if name:
                prn_meds.append(name)
                if entry["above_bnf_radio"].isChecked():
                    above_bnf_meds.append(name)

        # Build output text
        parts = []

        if regular_meds:
            parts.append(f"Regular: {', '.join(regular_meds)}")

        if prn_meds:
            parts.append(f"PRN: {', '.join(prn_meds)}")

        if parts:
            text = "; ".join(parts) + "."
            if above_bnf_meds:
                text += f"\n\nAll medication at BNF doses except {', '.join(above_bnf_meds)}."
            else:
                text += "\n\nAll medication at BNF doses."
            self.treatment_desc.setPlainText(text)
        else:
            self.treatment_desc.setPlainText("")

    def _build_form(self):
        # Clinician Details
        frame1 = self._create_section_frame("Approved Clinician / Medical Practitioner")
        layout1 = frame1.layout()

        # Clinician type radio buttons
        type_row = QHBoxLayout()
        type_row.setSpacing(16)
        type_lbl = QLabel("I am:")
        type_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        type_row.addWidget(type_lbl)

        self.clinician_type_group = QButtonGroup(self)
        self.approved_clinician_radio = QRadioButton("the approved clinician in charge of the treatment described below")
        self.approved_clinician_radio.setChecked(True)
        self.approved_clinician_radio.setStyleSheet("font-size: 12px;")
        self.soad_radio = QRadioButton("a registered medical practitioner appointed for the purposes of Part 4 of the Act (a SOAD)")
        self.soad_radio.setStyleSheet("font-size: 12px;")
        self.clinician_type_group.addButton(self.approved_clinician_radio)
        self.clinician_type_group.addButton(self.soad_radio)
        type_row.addWidget(self.approved_clinician_radio)
        type_row.addWidget(self.soad_radio)
        type_row.addStretch()
        layout1.addLayout(type_row)

        row1 = QHBoxLayout()
        row1.setSpacing(12)
        self.clinician_name = self._create_line_edit("Full name")
        row1.addWidget(self.clinician_name, 1)
        self.clinician_address = self._create_line_edit("Address")
        row1.addWidget(self.clinician_address, 2)
        self.clinician_email = self._create_line_edit("Email")
        row1.addWidget(self.clinician_email, 1)
        layout1.addLayout(row1)

        self.form_layout.addWidget(frame1)

        # Patient Details
        frame2 = self._create_section_frame("Patient Details")
        layout2 = frame2.layout()

        row2 = QHBoxLayout()
        row2.setSpacing(12)
        self.patient_name = self._create_line_edit("Patient full name")
        row2.addWidget(self.patient_name, 1)
        self.patient_address = self._create_line_edit("Patient address")
        row2.addWidget(self.patient_address, 2)
        layout2.addLayout(row2)

        self.form_layout.addWidget(frame2)

        # Treatment Description - split layout
        frame3 = self._create_section_frame("Treatment Description")
        layout3 = frame3.layout()

        info = QLabel("The patient is capable of understanding the nature, purpose and likely effects of the following treatment:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 13px; color: #374151;")
        layout3.addWidget(info)

        # Splitter: left = output text, right = medication inputs
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setStyleSheet("QSplitter::handle { background: #e5e7eb; width: 3px; }")

        # Left side - output text area
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 8, 0)
        self.treatment_desc = self._create_text_edit("Treatment summary will appear here...", 200)
        self.treatment_desc.setReadOnly(True)
        left_layout.addWidget(self.treatment_desc)
        splitter.addWidget(left_widget)

        # Right side - medication inputs
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(8, 0, 0, 0)
        right_layout.setSpacing(12)

        # Initialize medication lists
        self._regular_meds = []
        self._prn_meds = []

        # Regular Medications section
        regular_group = QGroupBox("Regular Medications")
        regular_group.setStyleSheet("""
            QGroupBox {
                font-weight: 600;
                font-size: 13px;
                color: #059669;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                margin-top: 8px;
                padding-top: 8px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
        """)
        regular_layout = QVBoxLayout(regular_group)
        regular_layout.setSpacing(6)

        self.regular_meds_container = QVBoxLayout()
        self.regular_meds_container.setSpacing(4)
        regular_layout.addLayout(self.regular_meds_container)

        add_regular_btn = QPushButton("+ Add Regular Med")
        add_regular_btn.setStyleSheet("""
            QPushButton {
                background: #e5e7eb;
                color: #374151;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover { background: #d1d5db; }
        """)
        add_regular_btn.clicked.connect(lambda: self._add_medication_entry("regular"))
        regular_layout.addWidget(add_regular_btn)
        right_layout.addWidget(regular_group)

        # PRN Medications section
        prn_group = QGroupBox("PRN Medications")
        prn_group.setStyleSheet("""
            QGroupBox {
                font-weight: 600;
                font-size: 13px;
                color: #059669;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                margin-top: 8px;
                padding-top: 8px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px;
            }
        """)
        prn_layout = QVBoxLayout(prn_group)
        prn_layout.setSpacing(6)

        self.prn_meds_container = QVBoxLayout()
        self.prn_meds_container.setSpacing(4)
        prn_layout.addLayout(self.prn_meds_container)

        add_prn_btn = QPushButton("+ Add PRN Med")
        add_prn_btn.setStyleSheet("""
            QPushButton {
                background: #e5e7eb;
                color: #374151;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover { background: #d1d5db; }
        """)
        add_prn_btn.clicked.connect(lambda: self._add_medication_entry("prn"))
        prn_layout.addWidget(add_prn_btn)
        right_layout.addWidget(prn_group)

        right_layout.addStretch()
        splitter.addWidget(right_widget)
        splitter.setSizes([400, 350])

        layout3.addWidget(splitter)

        consent_lbl = QLabel("AND has consented to that treatment.")
        consent_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #059669; margin-top: 8px;")
        layout3.addWidget(consent_lbl)

        self.form_layout.addWidget(frame3)

        # Add initial medication entries
        self._add_medication_entry("regular")
        self._add_medication_entry("prn")

        # Signature
        frame4 = self._create_section_frame("Signature")
        layout4 = frame4.layout()

        sig_row = QHBoxLayout()
        sig_row.setSpacing(12)
        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        sig_row.addWidget(self.sig_date)
        sig_row.addStretch()
        layout4.addLayout(sig_row)

        self.form_layout.addWidget(frame4)

    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.clinician_name.clear()
            self.clinician_address.clear()
            self.clinician_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.treatment_desc.clear()
            self.sig_date.setDate(QDate.currentDate())
            # Clear medication entries
            for entry in self._regular_meds:
                entry["name"].setCurrentText("")
                entry["bnf_radio"].setChecked(True)
            for entry in self._prn_meds:
                entry["name"].setCurrentText("")
                entry["bnf_radio"].setChecked(True)
            # Reset clinician type
            self.approved_clinician_radio.setChecked(True)

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form T2",
            f"Form_T2_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_T2_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form T2 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def add_strikethrough(run):
                """Add strikethrough to a run."""
                rPr = run._element.get_or_add_rPr()
                strike = OxmlElement('w:strike')
                strike.set(qn('w:val'), 'true')
                rPr.append(strike)

            paragraphs = doc.paragraphs

            # Clinician type (para 4) - strikethrough unselected option
            para4_text = paragraphs[4].text
            if "approved clinician" in para4_text and "registered medical practitioner" in para4_text:
                # Clear existing runs
                for run in paragraphs[4].runs:
                    run.text = ""

                if self.approved_clinician_radio.isChecked():
                    # Keep approved clinician, strikethrough SOAD
                    run1 = paragraphs[4].add_run("the approved clinician in charge of the treatment described below")
                    run2 = paragraphs[4].add_run("/")
                    add_strikethrough(run2)
                    run3 = paragraphs[4].add_run("a registered medical practitioner appointed for the purposes of Part 4 of the Act (a SOAD)")
                    add_strikethrough(run3)
                else:
                    # Strikethrough approved clinician, keep SOAD
                    run1 = paragraphs[4].add_run("the approved clinician in charge of the treatment described below")
                    add_strikethrough(run1)
                    run2 = paragraphs[4].add_run("/")
                    add_strikethrough(run2)
                    run3 = paragraphs[4].add_run("a registered medical practitioner appointed for the purposes of Part 4 of the Act (a SOAD)")

            # Clinician details (para 3)
            clinician_text = self.clinician_name.text()
            if self.clinician_address.text():
                clinician_text += ", " + self.clinician_address.text()
            if self.clinician_email.text():
                clinician_text += ", Email: " + self.clinician_email.text()
            if clinician_text.strip():
                set_para_text(paragraphs[3], clinician_text)
                highlight_yellow(paragraphs[3])

            # Patient details (para 6)
            patient_text = self.patient_name.text()
            if self.patient_address.text():
                patient_text += ", " + self.patient_address.text()
            if patient_text.strip():
                set_para_text(paragraphs[6], patient_text)
                highlight_yellow(paragraphs[6])

            # Treatment description (para 8-10)
            if self.treatment_desc.toPlainText().strip():
                set_para_text(paragraphs[8], self.treatment_desc.toPlainText())
                highlight_yellow(paragraphs[8])

            # Signature date (para 14)
            sig_date = self.sig_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[14], f"Signed                                              Date {sig_date}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form T2 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    def get_state(self) -> dict:
        # Collect medication data
        regular_meds = []
        for entry in self._regular_meds:
            name = entry["name"].currentText().strip()
            if name:
                regular_meds.append({
                    "name": name,
                    "above_bnf": entry["above_bnf_radio"].isChecked()
                })

        prn_meds = []
        for entry in self._prn_meds:
            name = entry["name"].currentText().strip()
            if name:
                prn_meds.append({
                    "name": name,
                    "above_bnf": entry["above_bnf_radio"].isChecked()
                })

        return {
            "clinician_type": "soad" if self.soad_radio.isChecked() else "approved_clinician",
            "clinician_name": self.clinician_name.text(),
            "clinician_address": self.clinician_address.text(),
            "clinician_email": self.clinician_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "treatment_desc": self.treatment_desc.toPlainText(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
            "regular_meds": regular_meds,
            "prn_meds": prn_meds,
        }

    def set_state(self, state: dict):
        if not state:
            return
        # Restore clinician type
        if state.get("clinician_type") == "soad":
            self.soad_radio.setChecked(True)
        else:
            self.approved_clinician_radio.setChecked(True)
        self.clinician_name.setText(state.get("clinician_name", ""))
        self.clinician_address.setText(state.get("clinician_address", ""))
        self.clinician_email.setText(state.get("clinician_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.treatment_desc.setPlainText(state.get("treatment_desc", ""))
        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))

        # Restore regular medications
        regular_meds = state.get("regular_meds", [])
        for i, med in enumerate(regular_meds):
            if i >= len(self._regular_meds):
                self._add_medication_entry("regular")
            entry = self._regular_meds[i]
            entry["name"].setCurrentText(med.get("name", ""))
            if med.get("above_bnf"):
                entry["above_bnf_radio"].setChecked(True)
            else:
                entry["bnf_radio"].setChecked(True)

        # Restore PRN medications
        prn_meds = state.get("prn_meds", [])
        for i, med in enumerate(prn_meds):
            if i >= len(self._prn_meds):
                self._add_medication_entry("prn")
            entry = self._prn_meds[i]
            entry["name"].setCurrentText(med.get("name", ""))
            if med.get("above_bnf"):
                entry["above_bnf_radio"].setChecked(True)
            else:
                entry["bnf_radio"].setChecked(True)

        self._update_treatment_text()
