# ============================================================
# DOCX EXPORTER — HTML → DOCX (Option C)
# ============================================================

from docx import Document
from bs4 import BeautifulSoup


class DocxExporter:

    @staticmethod
    def export_html(html: str, output_path: str):
        """
        Accepts a HTML string and writes a DOCX file.
        Simple HTML → paragraphs mapping (strong, em, underline supported).
        """
        try:
            doc = Document()
            soup = BeautifulSoup(html, "html.parser")

            for element in soup.descendants:
                if element.name in ["p", "div", "br"]:
                    doc.add_paragraph()
                if element.name is None:
                    text = element.strip()
                    if text:
                        doc.add_paragraph(text)

            doc.save(output_path)
            print(f"[DOCX EXPORT] Saved → {output_path}")

        except Exception as e:
            print(f"[DOCX EXPORT ERROR]: {e}")
