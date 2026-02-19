# ================================================================
#  MOJ LEAVE FORM PAGE — Leave Application for Restricted Patient
#  Mental Health Casework Section (MHCS) Leave Application
#  Based exactly on MHCS_Leave_Application_Form.docx structure
#  Restructured with card-based layout (like tribunal report)
# ================================================================

from __future__ import annotations
import re
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QEvent
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QGridLayout, QRadioButton, QButtonGroup, QComboBox, QSlider,
    QSplitter, QSpinBox, QSizePolicy, QStackedWidget
)
from background_history_popup import CollapsibleSection, ResizableSection


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# PRONOUN HELPER
# ================================================================
def pronouns_from_gender(g):
    """Return pronoun dictionary based on gender."""
    g = (g or "").lower().strip()
    if g == "male":
        return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "his", "have": "has", "are": "is", "do": "does"}
    if g == "female":
        return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "her", "have": "has", "are": "is", "do": "does"}
    return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "their", "have": "have", "are": "are", "do": "do"}


# ================================================================
# ZOOM HELPER FUNCTION
# ================================================================

def create_zoom_row(text_edit: QTextEdit, base_size: int = 12) -> QHBoxLayout:
    """Create a zoom controls row for any QTextEdit."""
    zoom_row = QHBoxLayout()
    zoom_row.setSpacing(2)
    zoom_row.addStretch()

    text_edit._font_size = base_size

    zoom_out_btn = QPushButton("−")
    zoom_out_btn.setFixedSize(16, 16)
    zoom_out_btn.setCursor(Qt.CursorShape.PointingHandCursor)
    zoom_out_btn.setToolTip("Decrease font size")
    zoom_out_btn.setStyleSheet("""
        QPushButton {
            background: #f3f4f6;
            color: #374151;
            border: 1px solid #d1d5db;
            border-radius: 3px;
            font-size: 11px;
            font-weight: bold;
        }
        QPushButton:hover { background: #e5e7eb; }
    """)
    zoom_row.addWidget(zoom_out_btn)

    zoom_in_btn = QPushButton("+")
    zoom_in_btn.setFixedSize(16, 16)
    zoom_in_btn.setCursor(Qt.CursorShape.PointingHandCursor)
    zoom_in_btn.setToolTip("Increase font size")
    zoom_in_btn.setStyleSheet("""
        QPushButton {
            background: #f3f4f6;
            color: #374151;
            border: 1px solid #d1d5db;
            border-radius: 3px;
            font-size: 11px;
            font-weight: bold;
        }
        QPushButton:hover { background: #e5e7eb; }
    """)
    zoom_row.addWidget(zoom_in_btn)

    current_style = text_edit.styleSheet()

    def zoom_in():
        text_edit._font_size = min(text_edit._font_size + 2, 28)
        new_style = re.sub(r'font-size:\s*\d+px', f'font-size: {text_edit._font_size}px', current_style)
        text_edit.setStyleSheet(new_style)

    def zoom_out():
        text_edit._font_size = max(text_edit._font_size - 2, 8)
        new_style = re.sub(r'font-size:\s*\d+px', f'font-size: {text_edit._font_size}px', current_style)
        text_edit.setStyleSheet(new_style)

    zoom_in_btn.clicked.connect(zoom_in)
    zoom_out_btn.clicked.connect(zoom_out)

    return zoom_row


# ================================================================
# MOJ LEAVE CARD WIDGET
# ================================================================

class MOJLeaveCardWidget(QFrame):
    """Card widget for MOJ Leave Form sections (similar to TribunalCardWidget)."""

    clicked = Signal(str)

    STYLE_NORMAL = """
        QFrame#mojCard {
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 12px;
        }
        QFrame#mojCard:hover {
            border: 1px solid #fca5a5;
            background: #fef2f2;
        }
    """
    STYLE_SELECTED = """
        QFrame#mojCard {
            background: #fef2f2;
            border: 2px solid #dc2626;
            border-radius: 12px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.title = title
        self.key = key
        self._selected = False
        self.setObjectName("mojCard")
        self.setStyleSheet(self.STYLE_NORMAL)
        self.setCursor(Qt.CursorShape.PointingHandCursor)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(8)

        # Title
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 14px;
            font-weight: 600;
            color: #1f2937;
        """)
        layout.addWidget(title_lbl)

        # Editor (preview/edit area)
        self.editor = QTextEdit()
        self.editor.setReadOnly(False)
        self._editor_height = 120
        self.editor.setFixedHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 12px;
                color: #374151;
            }
        """)
        editor_zoom = create_zoom_row(self.editor, base_size=12)
        layout.addLayout(editor_zoom)
        layout.addWidget(self.editor)

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(12)
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #dc2626;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        """Handle drag events on the expand bar."""
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(80, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        return self._selected


# ================================================================
# MOJ LEAVE FORM PAGE
# ================================================================

class MOJLeaveFormPage(QWidget):
    """Page for completing MOJ Leave Application for Restricted Patients."""

    go_back = Signal()

    # Sections based on MHCS Leave Application Form
    SECTIONS = [
        ("1. Patient Details", "patient_details"),
        ("2. Responsible Clinician Details", "rc_details"),
        ("3a. Type of Leave Requested", "leave_type"),
        ("3b. Documents Reviewed", "documents"),
        ("3c. Purpose of Leave", "purpose"),
        ("3d. Unescorted Overnight Leave", "overnight"),
        ("3e. Escorted Overnight Leave", "escorted_overnight"),
        ("3f. Compassionate Leave", "compassionate"),
        ("3g. Leave Report", "leave_report"),
        ("3h. Proposed Management", "procedures"),
        ("4a. Past Psychiatric History", "hospital_admissions"),
        ("4b. Index Offence and Forensic History", "index_offence"),
        ("4c. Current Mental Disorder", "mental_disorder"),
        ("4d. Attitude and Behaviour", "attitude_behaviour"),
        ("4e. Risk Factors", "risk_factors"),
        ("4f. Medication", "medication"),
        ("4g. Psychology", "psychology"),
        ("4h. Extremism", "extremism"),
        ("4i. Absconding", "absconding"),
        ("5. MAPPA", "mappa"),
        ("6. Victims", "victims"),
        ("7. Transferred Prisoners", "transferred_prisoners"),
        ("8. Fitness to Plead", "fitness_to_plead"),
        ("9. Additional Comments", "additional_comments"),
        ("Signature", "signature"),
        ("Annex A - Victim Liaison", "annex_a"),
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self.cards = {}
        self.popups = {}
        self.popup_previews = {}  # Preview labels for each popup
        self.popup_send_buttons = {}  # Send buttons for each popup
        self.popup_generators = {}  # Generator functions for each popup
        self._selected_card_key = None
        self._my_details = self._load_my_details()
        self._extracted_raw_notes = []
        self._extracted_categories = {}
        self._data_extractor_overlay = None
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        # Pre-fill RC details from database
        pass  # Will be handled in popup

    def _get_pronouns(self):
        """Get pronouns based on gender selection."""
        # Gender radio buttons are defined on self, not on the popup
        if hasattr(self, 'gender_male') and self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his", "self": "himself", "has": "has", "have": "has", "is": "is", "are": "is", "do": "does"}
        elif hasattr(self, 'gender_female') and self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her", "self": "herself", "has": "has", "have": "has", "is": "is", "are": "is", "do": "does"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their", "self": "themselves", "has": "have", "have": "have", "is": "are", "are": "are", "do": "do"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #991b1b; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("MOJ Leave Application — Restricted Patient")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Clear Form button
        clear_btn = QPushButton("Clear Form - Start New")
        clear_btn.setFixedSize(180, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #6b7280;
                color: white;
                font-size: 12px;
                font-weight: 600;
                border: none;
                border-radius: 6px;
            }
            QPushButton:hover { background: #4b5563; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        toolbar = QWidget()
        toolbar.setFixedHeight(50)
        toolbar.setStyleSheet("background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12);")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 6, 16, 6)
        tb_layout.setSpacing(12)

        # Export DOCX button
        export_btn = QPushButton("Export DOCX")
        export_btn.setFixedSize(120, 36)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        export_btn.clicked.connect(self._export_docx)
        tb_layout.addWidget(export_btn)

        # Import Data button
        import_btn = QPushButton("Import Data")
        import_btn.setFixedSize(120, 36)
        import_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        import_btn.setStyleSheet("""
            QPushButton {
                background: #2563eb;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QPushButton:hover { background: #1d4ed8; }
        """)
        import_btn.clicked.connect(self._import_data)
        tb_layout.addWidget(import_btn)

        # View Data button (purple, like ASR)
        view_data_btn = QPushButton("View Data")
        view_data_btn.setFixedSize(100, 36)
        view_data_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        view_data_btn.setStyleSheet("""
            QPushButton {
                background: #8b5cf6;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QPushButton:hover { background: #7c3aed; }
            QPushButton:pressed { background: #6d28d9; }
        """)
        view_data_btn.clicked.connect(self._view_data)
        tb_layout.addWidget(view_data_btn)

        tb_layout.addStretch()
        main_layout.addWidget(toolbar)

        # Content area with splitter
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards in scroll area
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("""
            QScrollArea {
                background: #f9fafb;
                border: none;
            }
        """)
        self.main_splitter.addWidget(self.cards_holder)

        self.cards_root = QWidget()
        self.cards_root.setStyleSheet("background: #f9fafb;")
        self.cards_layout = QVBoxLayout(self.cards_root)
        self.cards_layout.setContentsMargins(32, 24, 32, 24)
        self.cards_layout.setSpacing(16)
        self.cards_holder.setWidget(self.cards_root)

        # Right: Panel with popup stack
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(350)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame#editorPanel {
                background: white;
                border-left: 1px solid #e5e7eb;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QCheckBox, QRadioButton {
                background: transparent;
                border: none;
            }
        """)
        self.editor_panel.setObjectName("editorPanel")
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 450])

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 18px;
            font-weight: 700;
            color: #1f2937;
            background: white;
            padding: 8px 12px;
            border-radius: 8px;
            border: 1px solid #e5e7eb;
        """)
        panel_layout.addWidget(self.panel_title)

        # Popup stack
        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        # Create all cards
        self._create_cards()

    def _create_cards(self):
        """Create all section cards."""
        for title, key in self.SECTIONS:
            card = MOJLeaveCardWidget(title, key, parent=self.cards_root)
            card.clicked.connect(self._on_card_clicked)
            self.cards[key] = card
            self.cards_layout.addWidget(card)

        self.cards_layout.addStretch()

    def _on_card_clicked(self, key: str):
        """Handle card click - show appropriate popup."""
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        # Update card selection highlighting
        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)
        if key in self.cards:
            self.cards[key].setSelected(True)
        self._selected_card_key = key

        # Create popup if not exists
        if key not in self.popups:
            popup = self._create_popup(key)
            if popup:
                self.popups[key] = popup
                self.popup_stack.addWidget(popup)

        # Show popup
        if key in self.popups:
            self.popup_stack.setCurrentWidget(self.popups[key])

    def _create_popup(self, key: str) -> QWidget:
        """Create popup widget for the given section key."""
        if key == "patient_details":
            return self._create_patient_details_popup()
        elif key == "rc_details":
            return self._create_rc_details_popup()
        elif key == "leave_type":
            return self._create_leave_type_popup()
        elif key == "documents":
            return self._create_documents_popup()
        elif key == "purpose":
            return self._create_purpose_popup()
        elif key == "overnight":
            return self._create_overnight_popup()
        elif key == "escorted_overnight":
            return self._create_escorted_overnight_popup()
        elif key == "compassionate":
            return self._create_compassionate_popup()
        elif key == "leave_report":
            return self._create_leave_report_popup()
        elif key == "procedures":
            return self._create_procedures_popup()
        elif key == "hospital_admissions":
            return self._create_hospital_admissions_popup()
        elif key == "index_offence":
            return self._create_index_offence_popup()
        elif key == "mental_disorder":
            return self._create_mental_disorder_popup()
        elif key == "attitude_behaviour":
            return self._create_attitude_behaviour_popup()
        elif key == "risk_factors":
            return self._create_risk_factors_popup()
        elif key == "medication":
            return self._create_medication_popup()
        elif key == "psychology":
            return self._create_psychology_popup()
        elif key == "extremism":
            return self._create_extremism_popup()
        elif key == "absconding":
            return self._create_absconding_popup()
        elif key == "mappa":
            return self._create_mappa_popup()
        elif key == "victims":
            return self._create_victims_popup()
        elif key == "transferred_prisoners":
            return self._create_transferred_prisoners_popup()
        elif key == "fitness_to_plead":
            return self._create_fitness_to_plead_popup()
        elif key == "additional_comments":
            return self._create_additional_comments_popup()
        elif key == "signature":
            return self._create_signature_popup()
        elif key == "annex_a":
            return self._create_annex_a_popup()
        return None

    # ================================================================
    # POPUP CONTAINER HELPER (Preview + Input layout)
    # ================================================================

    def _create_popup_container(self, key: str) -> tuple:
        """Create popup with preview section at top and input section below.
        Returns (main_widget, input_layout) for adding input fields."""
        main_widget = QWidget()
        main_widget.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Use QSplitter for resizable preview/input sections
        splitter = QSplitter(Qt.Orientation.Vertical)
        splitter.setHandleWidth(8)
        splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(0,0,0,0.05), stop:0.5 rgba(0,0,0,0.15), stop:1 rgba(0,0,0,0.05));
                margin: 2px 60px;
                border-radius: 3px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(153,27,27,0.2), stop:0.5 rgba(153,27,27,0.5), stop:1 rgba(153,27,27,0.2));
            }
        """)
        main_layout.addWidget(splitter)

        # === PREVIEW SECTION (at top) ===
        preview_widget = QWidget()
        preview_widget.setStyleSheet("background: white;")
        preview_widget_layout = QVBoxLayout(preview_widget)
        preview_widget_layout.setContentsMargins(0, 0, 0, 4)
        preview_widget_layout.setSpacing(0)

        preview_container = QWidget()
        preview_container.setStyleSheet("""
            QWidget {
                background: rgba(255,255,255,0.96);
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 8px;
            }
        """)
        preview_container_layout = QVBoxLayout(preview_container)
        preview_container_layout.setContentsMargins(12, 8, 12, 8)
        preview_container_layout.setSpacing(6)

        # Preview header with title and Send button
        preview_header = QHBoxLayout()
        preview_title = QLabel("Preview")
        preview_title.setStyleSheet("""
            QLabel {
                font-size: 13px;
                font-weight: 600;
                color: #991b1b;
                background: transparent;
                border: none;
            }
        """)
        preview_header.addWidget(preview_title)
        preview_header.addStretch()

        # Send to Card button
        send_btn = QPushButton("Send to Card")
        send_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        send_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                padding: 6px 14px;
                border-radius: 5px;
                font-size: 12px;
                font-weight: 600;
                border: none;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #6b1515; }
        """)
        preview_header.addWidget(send_btn)
        preview_container_layout.addLayout(preview_header)

        # Dark preview area
        preview_scroll = QScrollArea()
        preview_scroll.setWidgetResizable(True)
        preview_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        preview_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        preview_scroll.setStyleSheet("QScrollArea { background: #1e1e1e; border-radius: 6px; border: none; }")

        preview_label = QLabel("")
        preview_label.setWordWrap(True)
        preview_label.setAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
        preview_label.setStyleSheet("""
            QLabel {
                background: #1e1e1e;
                color: #eaeaea;
                padding: 10px;
                font-size: 12px;
                border: none;
            }
        """)
        preview_scroll.setWidget(preview_label)
        preview_container_layout.addWidget(preview_scroll, 1)
        preview_widget_layout.addWidget(preview_container)
        splitter.addWidget(preview_widget)

        # Store preview references
        self.popup_previews[key] = preview_label
        self.popup_send_buttons[key] = send_btn

        # === INPUT SECTION (below) ===
        input_scroll = QScrollArea()
        input_scroll.setWidgetResizable(True)
        input_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        input_scroll.setStyleSheet("QScrollArea { background: white; border: none; }")

        input_content = QWidget()
        input_content.setStyleSheet("background: white;")
        input_layout = QVBoxLayout(input_content)
        input_layout.setContentsMargins(4, 8, 4, 8)
        input_layout.setSpacing(10)

        input_scroll.setWidget(input_content)
        splitter.addWidget(input_scroll)

        # Set initial sizes (preview: 140px, input: rest)
        splitter.setSizes([140, 400])
        splitter.setStretchFactor(0, 0)
        splitter.setStretchFactor(1, 1)

        return main_widget, input_layout

    def _setup_popup_preview(self, key: str, generate_func):
        """Connect the preview send button to the generator function and set up live preview."""
        self.popup_generators[key] = generate_func

        # Connect send button
        if key in self.popup_send_buttons:
            self.popup_send_buttons[key].clicked.connect(
                lambda checked, k=key: self._send_to_card(k)
            )

        # Initial preview update
        self._update_preview(key)

    def _update_preview(self, key: str):
        """Update the preview label for a popup with generated text."""
        if key in self.popup_previews and key in self.popup_generators:
            try:
                text = self.popup_generators[key]()
                self.popup_previews[key].setText(text if text else "(No content yet - fill in fields below)")
            except Exception as e:
                self.popup_previews[key].setText(f"(Preview error: {e})")

    def _send_to_card(self, key: str):
        """Send generated text to the card's editor."""
        if key in self.cards and key in self.popup_generators:
            text = self.popup_generators[key]()
            if text:
                current = self.cards[key].editor.toPlainText()
                if current:
                    self.cards[key].editor.setPlainText(current + "\n\n" + text)
                else:
                    self.cards[key].editor.setPlainText(text)

    def _connect_preview_updates(self, key: str, widgets: list):
        """Connect widgets to trigger preview updates when changed."""
        for widget in widgets:
            if isinstance(widget, QLineEdit):
                widget.textChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QTextEdit):
                widget.textChanged.connect(lambda k=key: self._update_preview(k))
            elif isinstance(widget, QComboBox):
                widget.currentTextChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QCheckBox):
                widget.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QRadioButton):
                widget.toggled.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QDateEdit):
                widget.dateChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QSlider):
                widget.valueChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QSpinBox):
                widget.valueChanged.connect(lambda _, k=key: self._update_preview(k))

    # ================================================================
    # POPUP CREATORS
    # ================================================================

    def _create_patient_details_popup(self) -> QWidget:
        """Create patient details popup with preview panel and form fields."""
        key = "patient_details"
        popup, layout = self._create_popup_container(key)

        # Input field styling
        field_style = "background: white; padding: 8px; border: 1px solid #d1d5db; border-radius: 6px;"

        # Patient Name
        layout.addWidget(QLabel("Patient Name:"))
        self.patient_name = QLineEdit()
        self.patient_name.setPlaceholderText("Full name")
        self.patient_name.setStyleSheet(field_style)
        layout.addWidget(self.patient_name)

        # DOB
        layout.addWidget(QLabel("Date of Birth:"))
        self.patient_dob = QDateEdit()
        self.patient_dob.setCalendarPopup(True)
        self.patient_dob.setDisplayFormat("dd/MM/yyyy")
        self.patient_dob.setStyleSheet(field_style)
        layout.addWidget(self.patient_dob)

        # Gender
        layout.addWidget(QLabel("Gender:"))
        gender_row = QHBoxLayout()
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("Male")
        self.gender_female = QRadioButton("Female")
        self.gender_other = QRadioButton("Other")
        self.gender_group.addButton(self.gender_male)
        self.gender_group.addButton(self.gender_female)
        self.gender_group.addButton(self.gender_other)
        gender_row.addWidget(self.gender_male)
        gender_row.addWidget(self.gender_female)
        gender_row.addWidget(self.gender_other)
        gender_row.addStretch()
        layout.addLayout(gender_row)

        # Hospital Number
        layout.addWidget(QLabel("Hospital Number:"))
        self.hospital_number = QLineEdit()
        self.hospital_number.setPlaceholderText("Hospital/NHS number")
        self.hospital_number.setStyleSheet(field_style)
        layout.addWidget(self.hospital_number)

        # Hospital Name
        layout.addWidget(QLabel("Hospital Name:"))
        self.hospital_name = QLineEdit()
        self.hospital_name.setPlaceholderText("Name of hospital")
        self.hospital_name.setStyleSheet(field_style)
        layout.addWidget(self.hospital_name)

        # Ward
        layout.addWidget(QLabel("Ward:"))
        self.ward = QLineEdit()
        self.ward.setPlaceholderText("Ward name")
        self.ward.setStyleSheet(field_style)
        layout.addWidget(self.ward)

        # MHA Section
        layout.addWidget(QLabel("Section of Mental Health Act:"))
        self.mha_section = QComboBox()
        self.mha_section.addItems(["", "37/41", "47/49", "48/49", "45A", "Other"])
        self.mha_section.setStyleSheet(field_style)
        self.mha_section.setEditable(True)
        self.mha_section.lineEdit().setPlaceholderText("Select or type section")
        layout.addWidget(self.mha_section)

        # Ministry of Justice Reference
        layout.addWidget(QLabel("Ministry of Justice Reference:"))
        self.moj_reference = QLineEdit()
        self.moj_reference.setPlaceholderText("MOJ reference number")
        self.moj_reference.setStyleSheet(field_style)
        layout.addWidget(self.moj_reference)

        layout.addStretch()

        # Setup preview generator and connect widgets
        def generate_patient_details():
            name = self.patient_name.text().strip()
            dob = self.patient_dob.date().toString("dd/MM/yyyy")
            gender = "Male" if self.gender_male.isChecked() else "Female" if self.gender_female.isChecked() else "Other" if self.gender_other.isChecked() else ""
            hosp_num = self.hospital_number.text().strip()
            hosp_name = self.hospital_name.text().strip()
            ward = self.ward.text().strip()
            section = self.mha_section.currentText().strip()
            moj_ref = self.moj_reference.text().strip()

            parts = []
            if name:
                parts.append(f"Patient: {name}")
            if dob and dob != QDate.currentDate().toString("dd/MM/yyyy"):
                parts.append(f"DOB: {dob}")
            if gender:
                parts.append(f"Gender: {gender}")
            if hosp_num:
                parts.append(f"Hospital No: {hosp_num}")
            if hosp_name:
                parts.append(f"Hospital: {hosp_name}")
            if ward:
                parts.append(f"Ward: {ward}")
            if section:
                parts.append(f"MHA Section: {section}")
            if moj_ref:
                parts.append(f"MOJ Ref: {moj_ref}")

            return "\n".join(parts) if parts else ""

        self._setup_popup_preview(key, generate_patient_details)
        self._connect_preview_updates(key, [
            self.patient_name, self.patient_dob, self.gender_male, self.gender_female,
            self.gender_other, self.hospital_number, self.hospital_name, self.ward,
            self.mha_section, self.moj_reference
        ])

        return popup

    def _create_rc_details_popup(self) -> QWidget:
        """Create RC details popup with preview panel."""
        key = "rc_details"
        popup, layout = self._create_popup_container(key)
        field_style = "background: white; padding: 8px; border: 1px solid #d1d5db; border-radius: 6px;"

        # RC Name
        layout.addWidget(QLabel("Responsible Clinician Name:"))
        self.rc_name = QLineEdit()
        self.rc_name.setPlaceholderText("Full name")
        self.rc_name.setStyleSheet(field_style)
        layout.addWidget(self.rc_name)

        # Pre-fill from database
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])

        # RC Email
        layout.addWidget(QLabel("Email:"))
        self.rc_email = QLineEdit()
        self.rc_email.setPlaceholderText("Email address")
        self.rc_email.setStyleSheet(field_style)
        layout.addWidget(self.rc_email)

        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

        # RC Phone
        layout.addWidget(QLabel("Phone:"))
        self.rc_phone = QLineEdit()
        self.rc_phone.setPlaceholderText("Phone number")
        self.rc_phone.setStyleSheet(field_style)
        layout.addWidget(self.rc_phone)

        layout.addStretch()

        # Setup preview
        def generate_rc_details():
            parts = []
            name = self.rc_name.text().strip()
            email = self.rc_email.text().strip()
            phone = self.rc_phone.text().strip()
            if name:
                parts.append(f"RC: {name}")
            if email:
                parts.append(f"Email: {email}")
            if phone:
                parts.append(f"Phone: {phone}")
            return "\n".join(parts) if parts else ""

        self._setup_popup_preview(key, generate_rc_details)
        self._connect_preview_updates(key, [self.rc_name, self.rc_email, self.rc_phone])

        return popup

    def _create_leave_type_popup(self) -> QWidget:
        """Create leave type popup with preview panel."""
        key = "leave_type"
        popup, layout = self._create_popup_container(key)

        layout.addWidget(QLabel("Select type(s) of leave requested:"))

        # Compassionate leave checkbox (always visible)
        self.compassionate_cb = QCheckBox("Compassionate/Emergency Leave")
        layout.addWidget(self.compassionate_cb)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("background: #e5e7eb;")
        layout.addWidget(sep)

        # Escorted / Unescorted radio buttons
        layout.addWidget(QLabel("Escort Status:"))
        self.escort_group = QButtonGroup(self)
        escort_row = QHBoxLayout()

        self.escorted_radio = QRadioButton("Escorted")
        self.unescorted_radio = QRadioButton("Unescorted")
        self.escort_group.addButton(self.escorted_radio)
        self.escort_group.addButton(self.unescorted_radio)
        escort_row.addWidget(self.escorted_radio)
        escort_row.addWidget(self.unescorted_radio)
        escort_row.addStretch()
        layout.addLayout(escort_row)

        # Leave options frame (shown when escorted/unescorted selected)
        self.leave_options_frame = QFrame()
        self.leave_options_frame.setStyleSheet("""
            QFrame {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding: 8px;
            }
        """)
        options_layout = QVBoxLayout(self.leave_options_frame)
        options_layout.setContentsMargins(12, 12, 12, 12)
        options_layout.setSpacing(8)

        options_layout.addWidget(QLabel("Leave Type:"))

        self.leave_type_checkboxes = {}
        leave_options = [
            ("ground", "Ground Leave"),
            ("community", "Community Leave"),
            ("overnight", "Overnight Leave"),
        ]

        for opt_key, label in leave_options:
            cb = QCheckBox(label)
            self.leave_type_checkboxes[opt_key] = cb
            options_layout.addWidget(cb)
            cb.stateChanged.connect(lambda _, k=key: self._update_preview(k))

        self.leave_options_frame.setVisible(False)
        layout.addWidget(self.leave_options_frame)

        # Connect radio buttons to show/hide options and update preview
        def on_escort_changed():
            show = self.escorted_radio.isChecked() or self.unescorted_radio.isChecked()
            self.leave_options_frame.setVisible(show)
            self._update_preview(key)

        self.escorted_radio.toggled.connect(on_escort_changed)
        self.unescorted_radio.toggled.connect(on_escort_changed)

        layout.addStretch()

        # Setup preview
        def generate_leave_type():
            parts = []
            if self.compassionate_cb.isChecked():
                parts.append("• Compassionate/Emergency Leave")
            escort = "Escorted" if self.escorted_radio.isChecked() else "Unescorted" if self.unescorted_radio.isChecked() else ""
            if escort:
                parts.append(f"• {escort}")
            for opt_key, label in leave_options:
                if self.leave_type_checkboxes[opt_key].isChecked():
                    parts.append(f"  - {label}")
            return "\n".join(parts) if parts else ""

        self._setup_popup_preview(key, generate_leave_type)
        self._connect_preview_updates(key, [self.compassionate_cb, self.escorted_radio, self.unescorted_radio])

        return popup

    def _create_documents_popup(self) -> QWidget:
        """Create documents reviewed popup with preview panel."""
        key = "documents"
        popup, layout = self._create_popup_container(key)

        layout.addWidget(QLabel("Documents reviewed in preparation for this application:"))

        self.documents_checkboxes = {}
        documents_list = [
            ("cpa_minutes", "CPA Minutes"),
            ("psychology_reports", "Psychology Reports"),
            ("hcr20", "HCR-20"),
            ("sara", "SARA"),
            ("other_risk_tools", "Other Risk Assessment Tools"),
            ("previous_reports", "Previous Reports"),
            ("current_reports", "Current Reports"),
            ("previous_notes", "Previous Notes"),
            ("current_notes", "Current Notes"),
        ]

        widgets_to_connect = []
        for doc_key, label in documents_list:
            cb = QCheckBox(label)
            self.documents_checkboxes[doc_key] = cb
            layout.addWidget(cb)
            widgets_to_connect.append(cb)

        # Other documents text field
        layout.addWidget(QLabel("Other documents (specify):"))
        self.other_documents = QLineEdit()
        self.other_documents.setPlaceholderText("Any other documents reviewed...")
        self.other_documents.setStyleSheet("background: white; padding: 8px; border: 1px solid #d1d5db; border-radius: 6px;")
        layout.addWidget(self.other_documents)
        widgets_to_connect.append(self.other_documents)

        layout.addStretch()

        # Setup preview
        def generate_documents():
            parts = []
            for doc_key, label in documents_list:
                if self.documents_checkboxes[doc_key].isChecked():
                    parts.append(f"• {label}")
            other = self.other_documents.text().strip()
            if other:
                parts.append(f"• Other: {other}")
            return "\n".join(parts) if parts else ""

        self._setup_popup_preview(key, generate_documents)
        self._connect_preview_updates(key, widgets_to_connect)

        return popup

    def _create_purpose_popup(self) -> QWidget:
        """Create purpose of leave popup with expandable preview and input sections."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { border: none; background: white; } QWidget { background: white; }")

        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        # === PREVIEW SECTION (at top, resizable) ===
        self.purpose_preview_section = ResizableSection()
        self.purpose_preview_section.set_content_height(120)

        preview_container = QFrame()
        preview_container.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
            }
        """)
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 8, 12, 8)
        preview_layout.setSpacing(4)

        preview_header_row = QHBoxLayout()
        preview_header_row.setContentsMargins(0, 0, 0, 0)

        preview_header = QLabel("Preview")
        preview_header.setStyleSheet("font-size: 13px; font-weight: 600; color: #991b1b;")
        preview_header_row.addWidget(preview_header)
        preview_header_row.addStretch()

        self.purpose_send_btn = QPushButton("Send to Card")
        self.purpose_send_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 5px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        self.purpose_send_btn.clicked.connect(lambda: self._send_to_card("purpose"))
        preview_header_row.addWidget(self.purpose_send_btn)

        preview_layout.addLayout(preview_header_row)

        self.purpose_preview = QTextEdit()
        self.purpose_preview.setReadOnly(True)
        self.purpose_preview.setStyleSheet("""
            QTextEdit {
                background: #1f2937;
                border: none;
                border-radius: 4px;
                padding: 8px;
                font-size: 12px;
                color: white;
            }
        """)
        preview_layout.addWidget(self.purpose_preview, 1)

        self.purpose_preview_section.set_content(preview_container)
        layout.addWidget(self.purpose_preview_section)

        # === INPUT SECTION (collapsible) ===
        self.purpose_input_section = CollapsibleSection("Input Fields")
        self.purpose_input_section.set_content_height(380)

        input_container = QFrame()
        input_container.setStyleSheet("QFrame { background: transparent; }")
        input_layout = QVBoxLayout(input_container)
        input_layout.setContentsMargins(8, 8, 8, 8)
        input_layout.setSpacing(10)

        # 1. PURPOSE OF LEAVE
        input_layout.addWidget(QLabel("<b>1. Purpose of Leave (therapeutic benefit):</b>"))

        self.purpose_group = QButtonGroup(self)
        self.purpose_radios = {}
        purpose_options = [
            ("starting", "Starting meaningful testing"),
            ("continuing", "Continuing previous leave"),
            ("unescorted", "Move to unescorted leave"),
            ("rehabilitation", "Rehabilitation process"),
        ]

        for key, label in purpose_options:
            radio = QRadioButton(label)
            radio.toggled.connect(self._update_purpose_preview)
            self.purpose_group.addButton(radio)
            self.purpose_radios[key] = radio
            input_layout.addWidget(radio)

        # 2. LOCATION OF LEAVE
        input_layout.addWidget(QLabel("<b>2. Location of Leave:</b>"))

        self.location_checkboxes = {}
        location_options = [
            ("ground", "Ground (hospital grounds)"),
            ("local", "Local (nearby area)"),
            ("community", "Community (wider area)"),
            ("family", "Family (visit family home)"),
        ]

        for key, label in location_options:
            cb = QCheckBox(label)
            cb.stateChanged.connect(self._update_purpose_preview)
            self.location_checkboxes[key] = cb
            input_layout.addWidget(cb)

        # Exclusion zone radios
        input_layout.addWidget(QLabel("Proximity to exclusion zone:"))
        self.exclusion_group = QButtonGroup(self)
        exclusion_row = QHBoxLayout()

        self.exclusion_yes = QRadioButton("Yes")
        self.exclusion_no = QRadioButton("No")
        self.exclusion_na = QRadioButton("N/A")

        self.exclusion_group.addButton(self.exclusion_yes)
        self.exclusion_group.addButton(self.exclusion_no)
        self.exclusion_group.addButton(self.exclusion_na)

        self.exclusion_yes.toggled.connect(self._update_purpose_preview)
        self.exclusion_no.toggled.connect(self._update_purpose_preview)
        self.exclusion_na.toggled.connect(self._update_purpose_preview)

        exclusion_row.addWidget(self.exclusion_yes)
        exclusion_row.addWidget(self.exclusion_no)
        exclusion_row.addWidget(self.exclusion_na)
        exclusion_row.addStretch()
        input_layout.addLayout(exclusion_row)

        # 3. DISCHARGE PLAN
        input_layout.addWidget(QLabel("<b>3. Discharge Planning Status:</b>"))

        self.discharge_options = ["Not started", "Early stages", "In progress", "Almost completed", "Completed"]
        self.discharge_label = QLabel(self.discharge_options[0])
        self.discharge_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        input_layout.addWidget(self.discharge_label)

        self.discharge_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.discharge_slider.setRange(0, len(self.discharge_options) - 1)
        self.discharge_slider.setValue(0)
        self.discharge_slider.setStyleSheet("""
            QSlider::groove:horizontal {
                border: 1px solid #d1d5db;
                height: 8px;
                background: #f3f4f6;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background: #991b1b;
                border: none;
                width: 18px;
                margin: -5px 0;
                border-radius: 9px;
            }
            QSlider::sub-page:horizontal {
                background: #fca5a5;
                border-radius: 4px;
            }
        """)
        self.discharge_slider.valueChanged.connect(
            lambda v: (self.discharge_label.setText(self.discharge_options[v]), self._update_purpose_preview())
        )
        input_layout.addWidget(self.discharge_slider)

        self.purpose_input_section.set_content(input_container)
        layout.addWidget(self.purpose_input_section)

        layout.addStretch()
        popup.setWidget(content)
        return popup

    def _get_patient_gender(self) -> str:
        """Get patient gender from the form, default to empty (they/them)."""
        if hasattr(self, 'gender_male') and self.gender_male.isChecked():
            return "male"
        elif hasattr(self, 'gender_female') and self.gender_female.isChecked():
            return "female"
        return ""  # Default to they/them

    def _update_purpose_preview(self):
        """Update the purpose preview with generated narrative."""
        if not hasattr(self, 'purpose_preview'):
            return

        # Get pronouns based on patient gender
        p = pronouns_from_gender(self._get_patient_gender())
        parts = []

        # 1. Purpose narrative
        if hasattr(self, 'purpose_radios'):
            if self.purpose_radios.get("starting") and self.purpose_radios["starting"].isChecked():
                parts.append("The purpose of leave is starting meaningful testing in the community looking toward progression to possible unescorted leave.")
            elif self.purpose_radios.get("continuing") and self.purpose_radios["continuing"].isChecked():
                parts.append(f"{p['subj']} {p['have']} already had some leave granted so the aim would be to continue to build on this.")
            elif self.purpose_radios.get("unescorted") and self.purpose_radios["unescorted"].isChecked():
                parts.append("The aim of leave is to move from escorted to unescorted to allow further independence and rehabilitation.")
            elif self.purpose_radios.get("rehabilitation") and self.purpose_radios["rehabilitation"].isChecked():
                parts.append(f"The leave is to build on {p['pos']} rehabilitation process.")

        # 2. Location narrative
        if hasattr(self, 'location_checkboxes'):
            locations = []
            if self.location_checkboxes.get("ground") and self.location_checkboxes["ground"].isChecked():
                locations.append("the hospital grounds")
            if self.location_checkboxes.get("local") and self.location_checkboxes["local"].isChecked():
                locations.append("the local area")
            if self.location_checkboxes.get("community") and self.location_checkboxes["community"].isChecked():
                locations.append("the wider community")
            if self.location_checkboxes.get("family") and self.location_checkboxes["family"].isChecked():
                locations.append("family residence")

            if locations:
                if len(locations) == 1:
                    parts.append(f"Leave would take place within {locations[0]}.")
                else:
                    loc_text = ", ".join(locations[:-1]) + " and " + locations[-1]
                    parts.append(f"Leave would take place within {loc_text}.")

        # Exclusion zone
        if hasattr(self, 'exclusion_yes') and self.exclusion_yes.isChecked():
            parts.append("The leave is close to/within the exclusion zone and this will be monitored closely by the team.")
        elif hasattr(self, 'exclusion_no') and self.exclusion_no.isChecked():
            parts.append("There are no concerns regarding the exclusion zone with this leave.")
        elif hasattr(self, 'exclusion_na') and self.exclusion_na.isChecked():
            parts.append("There is no exclusion zone with this patient.")

        # 3. Discharge planning narrative
        if hasattr(self, 'discharge_slider'):
            discharge_idx = self.discharge_slider.value()

            if discharge_idx == 0:  # Not started
                parts.append("Discharge planning has not yet commenced and leave would be an early step in this process.")
            elif discharge_idx == 1:  # Early stages
                parts.append("Discharge planning is in its early stages and leave would be an important step in building toward this.")
            elif discharge_idx == 2:  # In progress
                parts.append("Discharge planning is currently in progress and leave would be an important step in this.")
            elif discharge_idx == 3:  # Almost completed
                parts.append("Discharge planning is almost complete and leave would support the final stages of preparation.")
            elif discharge_idx == 4:  # Completed
                parts.append("Discharge planning is complete and leave forms part of the transition plan.")

        # Join with single space for flowing narrative
        self.purpose_preview.setPlainText(" ".join(parts))

    def _create_overnight_popup(self) -> QWidget:
        """Create unescorted overnight leave popup with comprehensive options."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { border: none; background: white; } QWidget { background: white; }")

        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        # === PREVIEW SECTION (compact, fixed height) ===
        preview_container = QFrame()
        preview_container.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; }")
        preview_container.setMaximumHeight(160)
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 8, 12, 8)
        preview_layout.setSpacing(4)

        preview_header_row = QHBoxLayout()
        preview_header_row.setContentsMargins(0, 0, 0, 0)

        preview_header = QLabel("Preview")
        preview_header.setStyleSheet("font-size: 13px; font-weight: 600; color: #991b1b;")
        preview_header_row.addWidget(preview_header)
        preview_header_row.addStretch()

        self.overnight_send_btn = QPushButton("Send to Card")
        self.overnight_send_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 5px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        self.overnight_send_btn.clicked.connect(lambda: self._send_to_card("overnight"))
        preview_header_row.addWidget(self.overnight_send_btn)

        preview_layout.addLayout(preview_header_row)

        self.overnight_preview = QTextEdit()
        self.overnight_preview.setReadOnly(True)
        self.overnight_preview.setMinimumHeight(60)
        self.overnight_preview.setMaximumHeight(100)
        self.overnight_preview.setStyleSheet("""
            QTextEdit { background: #1f2937; border: none; border-radius: 4px; padding: 8px; font-size: 12px; color: white; }
        """)
        preview_layout.addWidget(self.overnight_preview)

        layout.addWidget(preview_container)

        # === INPUT SECTION (direct, no collapsible wrapper) ===
        input_frame = QFrame()
        input_frame.setStyleSheet("QFrame { background: transparent; }")
        input_layout = QVBoxLayout(input_frame)
        input_layout.setContentsMargins(0, 0, 0, 0)
        input_layout.setSpacing(8)

        # === OVERNIGHT LEAVE - Yes/NA ===
        input_layout.addWidget(QLabel("<b>Unescorted overnight leave required:</b>"))
        self.overnight_main_group = QButtonGroup(self)
        overnight_main_row = QHBoxLayout()
        self.overnight_yes = QRadioButton("Yes")
        self.overnight_na = QRadioButton("N/A")
        self.overnight_main_group.addButton(self.overnight_yes)
        self.overnight_main_group.addButton(self.overnight_na)
        overnight_main_row.addWidget(self.overnight_yes)
        overnight_main_row.addWidget(self.overnight_na)
        overnight_main_row.addStretch()
        input_layout.addLayout(overnight_main_row)

        # === DETAILS CONTAINER (shown when Yes) ===
        self.overnight_details_container = QFrame()
        self.overnight_details_container.setStyleSheet("""
            QFrame { background: #f3f4f6; border: none; border-radius: 8px; }
            QLabel { background: #f3f4f6; border: none; }
            QRadioButton, QCheckBox { background: #f3f4f6; border: none; }
        """)
        details_layout = QVBoxLayout(self.overnight_details_container)
        details_layout.setContentsMargins(8, 8, 8, 8)
        details_layout.setSpacing(6)

        # === COLLAPSE 1: Accommodation ===
        self.accom_section = CollapsibleSection("1. Accommodation", start_collapsed=True)
        self.accom_section.set_content_height(220)
        accom_content = QFrame()
        accom_content.setStyleSheet("QFrame { background: #f3f4f6; } QLabel { background: #f3f4f6; }")
        accom_layout = QVBoxLayout(accom_content)
        accom_layout.setContentsMargins(8, 8, 8, 8)
        accom_layout.setSpacing(6)

        # Accommodation Type
        accom_layout.addWidget(QLabel("Type:"))
        self.accom_group = QButtonGroup(self)
        self.accom_24hr = QRadioButton("24 hour supported")
        self.accom_9to5 = QRadioButton("9-5 supported")
        self.accom_independent = QRadioButton("Independent")
        self.accom_family = QRadioButton("Family")
        self.accom_group.addButton(self.accom_24hr)
        self.accom_group.addButton(self.accom_9to5)
        self.accom_group.addButton(self.accom_independent)
        self.accom_group.addButton(self.accom_family)
        accom_type_row = QHBoxLayout()
        for rb in [self.accom_24hr, self.accom_9to5, self.accom_independent, self.accom_family]:
            rb.toggled.connect(self._update_overnight_preview)
            accom_type_row.addWidget(rb)
        accom_type_row.addStretch()
        accom_layout.addLayout(accom_type_row)

        # Address
        accom_layout.addWidget(QLabel("Address:"))
        self.overnight_address = QLineEdit()
        self.overnight_address.setPlaceholderText("Enter accommodation address...")
        self.overnight_address.setStyleSheet("background: white; color: #1f2937; padding: 8px; border: 1px solid #d1d5db; border-radius: 6px;")
        self.overnight_address.textChanged.connect(self._update_overnight_preview)
        accom_layout.addWidget(self.overnight_address)

        # Prior to Recall
        accom_layout.addWidget(QLabel("Address prior to recall:"))
        self.prior_recall_group = QButtonGroup(self)
        prior_row = QHBoxLayout()
        self.prior_recall_yes = QRadioButton("Yes")
        self.prior_recall_no = QRadioButton("No")
        self.prior_recall_group.addButton(self.prior_recall_yes)
        self.prior_recall_group.addButton(self.prior_recall_no)
        self.prior_recall_yes.toggled.connect(self._update_overnight_preview)
        self.prior_recall_no.toggled.connect(self._update_overnight_preview)
        prior_row.addWidget(self.prior_recall_yes)
        prior_row.addWidget(self.prior_recall_no)
        prior_row.addStretch()
        accom_layout.addLayout(prior_row)

        # Linked to Index Offence
        accom_layout.addWidget(QLabel("Linked to index offence:"))
        self.index_link_group = QButtonGroup(self)
        index_row = QHBoxLayout()
        self.index_link_yes = QRadioButton("Yes")
        self.index_link_no = QRadioButton("No")
        self.index_link_group.addButton(self.index_link_yes)
        self.index_link_group.addButton(self.index_link_no)
        self.index_link_yes.toggled.connect(self._update_overnight_preview)
        self.index_link_no.toggled.connect(self._update_overnight_preview)
        index_row.addWidget(self.index_link_yes)
        index_row.addWidget(self.index_link_no)
        index_row.addStretch()
        accom_layout.addLayout(index_row)

        self.accom_section.set_content(accom_content)
        details_layout.addWidget(self.accom_section)

        # === COLLAPSE 2: Support ===
        self.support_section = CollapsibleSection("2. Support", start_collapsed=True)
        self.support_section.set_content_height(100)
        support_content = QFrame()
        support_content.setStyleSheet("QFrame { background: #f3f4f6; } QLabel { background: #f3f4f6; } QCheckBox { background: #f3f4f6; }")
        support_layout = QVBoxLayout(support_content)
        support_layout.setContentsMargins(8, 8, 8, 8)
        support_layout.setSpacing(6)

        self.support_staff_cb = QCheckBox("Staff at the accommodation")
        self.support_cmht_cb = QCheckBox("CMHT")
        self.support_inpatient_cb = QCheckBox("Inpatient Team")
        self.support_family_cb = QCheckBox("Family")

        for cb in [self.support_staff_cb, self.support_cmht_cb, self.support_inpatient_cb, self.support_family_cb]:
            cb.stateChanged.connect(self._update_overnight_preview)
            support_layout.addWidget(cb)

        self.support_section.set_content(support_content)
        details_layout.addWidget(self.support_section)

        # === COLLAPSE 3: Number of Nights ===
        self.nights_section = CollapsibleSection("3. Number of Nights", start_collapsed=True)
        self.nights_section.set_content_height(100)
        nights_content = QFrame()
        nights_content.setStyleSheet("QFrame { background: #f3f4f6; } QLabel { background: #f3f4f6; }")
        nights_layout = QVBoxLayout(nights_content)
        nights_layout.setContentsMargins(8, 8, 8, 8)
        nights_layout.setSpacing(6)

        self.overnight_weeks_data = {}
        nights_row = QHBoxLayout()
        nights_row.addWidget(QLabel("Week:"))
        self.week_dropdown = QComboBox()
        self.week_dropdown.addItems([str(i) for i in range(1, 11)])
        self.week_dropdown.setStyleSheet("""
            QComboBox { background-color: white; color: black; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px; min-width: 50px; }
            QComboBox::drop-down { border: none; }
            QComboBox QAbstractItemView { background-color: white; color: black; selection-background-color: #e5e7eb; selection-color: black; }
        """)
        nights_row.addWidget(self.week_dropdown)
        nights_row.addWidget(QLabel("Nights:"))
        self.nights_dropdown = QComboBox()
        self.nights_dropdown.addItems([str(i) for i in range(1, 8)])
        self.nights_dropdown.setStyleSheet("""
            QComboBox { background-color: white; color: black; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px; min-width: 50px; }
            QComboBox::drop-down { border: none; }
            QComboBox QAbstractItemView { background-color: white; color: black; selection-background-color: #e5e7eb; selection-color: black; }
        """)
        nights_row.addWidget(self.nights_dropdown)
        nights_row.addStretch()
        nights_layout.addLayout(nights_row)

        # Auto-update on dropdown change
        def on_nights_changed():
            week = int(self.week_dropdown.currentText())
            nights = int(self.nights_dropdown.currentText())
            self.overnight_weeks_data[week] = nights
            # Update summary label
            summary_parts = [f"Wk {w}: {n}n" for w, n in sorted(self.overnight_weeks_data.items())]
            self.nights_summary_label.setText(", ".join(summary_parts))
            self._update_overnight_preview()

        self.week_dropdown.currentIndexChanged.connect(on_nights_changed)
        self.nights_dropdown.currentIndexChanged.connect(on_nights_changed)

        self.nights_summary_label = QLabel("")
        self.nights_summary_label.setStyleSheet("font-size: 11px; color: #374151; font-style: italic; background: #f3f4f6;")
        nights_layout.addWidget(self.nights_summary_label)

        self.nights_section.set_content(nights_content)
        details_layout.addWidget(self.nights_section)

        # === Discharge Address (not collapsible, always visible) ===
        discharge_frame = QFrame()
        discharge_frame.setStyleSheet("QFrame { background: #f3f4f6; } QLabel { background: #f3f4f6; }")
        discharge_layout_inner = QVBoxLayout(discharge_frame)
        discharge_layout_inner.setContentsMargins(0, 8, 0, 0)
        discharge_layout_inner.setSpacing(4)

        discharge_layout_inner.addWidget(QLabel("Leave to proposed discharge address:"))
        self.discharge_addr_group = QButtonGroup(self)
        discharge_row = QHBoxLayout()
        self.discharge_addr_yes = QRadioButton("Yes")
        self.discharge_addr_no = QRadioButton("No")
        self.discharge_addr_group.addButton(self.discharge_addr_yes)
        self.discharge_addr_group.addButton(self.discharge_addr_no)
        self.discharge_addr_yes.toggled.connect(self._update_overnight_preview)
        self.discharge_addr_no.toggled.connect(self._update_overnight_preview)
        discharge_row.addWidget(self.discharge_addr_yes)
        discharge_row.addWidget(self.discharge_addr_no)
        discharge_row.addStretch()
        discharge_layout_inner.addLayout(discharge_row)

        details_layout.addWidget(discharge_frame)

        self.overnight_details_container.setVisible(False)
        input_layout.addWidget(self.overnight_details_container)

        # Add input frame to main layout
        layout.addWidget(input_frame)

        # Connect Yes/NA to show/hide details
        def on_overnight_main_changed():
            if self.overnight_yes.isChecked():
                self.overnight_details_container.setVisible(True)
            else:
                self.overnight_details_container.setVisible(False)
            self._update_overnight_preview()

        self.overnight_yes.toggled.connect(on_overnight_main_changed)
        self.overnight_na.toggled.connect(on_overnight_main_changed)

        layout.addStretch()
        popup.setWidget(content)
        return popup

    def _update_overnight_preview(self):
        """Update the overnight leave preview with generated narrative."""
        if not hasattr(self, 'overnight_preview'):
            return

        p = pronouns_from_gender(self._get_patient_gender())
        parts = []

        # Check if N/A selected
        if hasattr(self, 'overnight_na') and self.overnight_na.isChecked():
            self.overnight_preview.setPlainText("Overnight leave is not applicable.")
            return

        if not (hasattr(self, 'overnight_yes') and self.overnight_yes.isChecked()):
            self.overnight_preview.setPlainText("")
            return

        # Accommodation type
        if hasattr(self, 'accom_24hr') and self.accom_24hr.isChecked():
            parts.append(f"{p['subj']} will be staying at 24 hour supported accommodation.")
        elif hasattr(self, 'accom_9to5') and self.accom_9to5.isChecked():
            parts.append(f"{p['subj']} will be staying at 9-5 supported accommodation.")
        elif hasattr(self, 'accom_independent') and self.accom_independent.isChecked():
            parts.append(f"{p['subj']} will be staying at independent accommodation.")
        elif hasattr(self, 'accom_family') and self.accom_family.isChecked():
            parts.append(f"{p['subj']} will be staying with family.")

        # Address
        if hasattr(self, 'overnight_address') and self.overnight_address.text().strip():
            parts.append(f"The address is {self.overnight_address.text().strip()}.")

        # Prior to recall
        if hasattr(self, 'prior_recall_yes') and self.prior_recall_yes.isChecked():
            parts.append(f"This was the address prior to recall and we believe it is a suitable disposal currently.")
        elif hasattr(self, 'prior_recall_no') and self.prior_recall_no.isChecked():
            parts.append(f"Since recall {p['pos']} needs have changed and {p['subj_l']} will be moved to a new address.")

        # Linked to index offence
        if hasattr(self, 'index_link_yes') and self.index_link_yes.isChecked():
            parts.append("The address is linked to the index offence and appropriate measures have been put in place to manage any ongoing concerns. The team will monitor the risk closely.")
        elif hasattr(self, 'index_link_no') and self.index_link_no.isChecked():
            parts.append("The address has no link to the index offence.")

        # Support
        support_sources = []
        if hasattr(self, 'support_staff_cb') and self.support_staff_cb.isChecked():
            support_sources.append("staff at the accommodation")
        if hasattr(self, 'support_cmht_cb') and self.support_cmht_cb.isChecked():
            support_sources.append("the Community Mental Health Team")
        if hasattr(self, 'support_inpatient_cb') and self.support_inpatient_cb.isChecked():
            support_sources.append("the inpatient team")
        if hasattr(self, 'support_family_cb') and self.support_family_cb.isChecked():
            support_sources.append("family members")

        if support_sources:
            if len(support_sources) == 1:
                parts.append(f"{p['subj']} will be supported by {support_sources[0]}.")
            elif len(support_sources) == 2:
                parts.append(f"{p['subj']} will be supported by {support_sources[0]} and {support_sources[1]}.")
            else:
                parts.append(f"{p['subj']} will be supported by {', '.join(support_sources[:-1])}, and {support_sources[-1]}.")

        # Number of nights
        if hasattr(self, 'overnight_weeks_data') and self.overnight_weeks_data:
            nights_parts = [f"Week {w}: {n} night{'s' if n > 1 else ''}" for w, n in sorted(self.overnight_weeks_data.items())]
            parts.append("Leave schedule: " + ". ".join(nights_parts) + ".")

        # Discharge address
        if hasattr(self, 'discharge_addr_yes') and self.discharge_addr_yes.isChecked():
            parts.append("I can confirm this leave is to the proposed discharge address.")
        elif hasattr(self, 'discharge_addr_no') and self.discharge_addr_no.isChecked():
            parts.append(f"This leave is not to the proposed discharge address but such testing is a necessary part of {p['pos']} rehabilitation.")

        self.overnight_preview.setPlainText(" ".join(parts))

    def _create_escorted_overnight_popup(self) -> QWidget:
        """Create escorted overnight leave popup with capacity/DoLS/discharge plan logic."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { border: none; background: white; } QWidget { background: white; }")

        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        # === PREVIEW SECTION (fixed height) ===
        preview_container = QFrame()
        preview_container.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; }")
        preview_container.setMaximumHeight(160)
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 8, 12, 8)
        preview_layout.setSpacing(4)

        preview_header_row = QHBoxLayout()
        preview_header_row.setContentsMargins(0, 0, 0, 0)

        preview_header = QLabel("Preview")
        preview_header.setStyleSheet("font-size: 13px; font-weight: 600; color: #991b1b;")
        preview_header_row.addWidget(preview_header)
        preview_header_row.addStretch()

        self.escorted_overnight_send_btn = QPushButton("Send to Card")
        self.escorted_overnight_send_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 5px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        self.escorted_overnight_send_btn.clicked.connect(lambda: self._send_to_card("escorted_overnight"))
        preview_header_row.addWidget(self.escorted_overnight_send_btn)

        preview_layout.addLayout(preview_header_row)

        self.escorted_overnight_preview = QTextEdit()
        self.escorted_overnight_preview.setReadOnly(True)
        self.escorted_overnight_preview.setMinimumHeight(60)
        self.escorted_overnight_preview.setMaximumHeight(100)
        self.escorted_overnight_preview.setStyleSheet("""
            QTextEdit { background: #1f2937; border: none; border-radius: 4px; padding: 8px; font-size: 12px; color: white; }
        """)
        preview_layout.addWidget(self.escorted_overnight_preview)

        layout.addWidget(preview_container)

        # === INPUT SECTION ===
        input_frame = QFrame()
        input_frame.setStyleSheet("QFrame { background: transparent; }")
        input_layout = QVBoxLayout(input_frame)
        input_layout.setContentsMargins(0, 0, 0, 0)
        input_layout.setSpacing(8)

        # === ESCORTED OVERNIGHT LEAVE - Yes/NA ===
        input_layout.addWidget(QLabel("<b>Escorted overnight leave required:</b>"))
        self.escorted_overnight_main_group = QButtonGroup(self)
        escorted_main_row = QHBoxLayout()
        self.escorted_overnight_yes = QRadioButton("Yes")
        self.escorted_overnight_na = QRadioButton("N/A")
        self.escorted_overnight_main_group.addButton(self.escorted_overnight_yes)
        self.escorted_overnight_main_group.addButton(self.escorted_overnight_na)
        escorted_main_row.addWidget(self.escorted_overnight_yes)
        escorted_main_row.addWidget(self.escorted_overnight_na)
        escorted_main_row.addStretch()
        input_layout.addLayout(escorted_main_row)

        # === DETAILS CONTAINER (shown when Yes) ===
        self.escorted_overnight_details_container = QFrame()
        self.escorted_overnight_details_container.setStyleSheet("QFrame { background: #f3f4f6; border-radius: 8px; } QLabel { background: #f3f4f6; border: none; } QRadioButton, QCheckBox { background: #f3f4f6; border: none; }")
        details_outer_layout = QVBoxLayout(self.escorted_overnight_details_container)
        details_outer_layout.setContentsMargins(12, 12, 12, 12)
        details_outer_layout.setSpacing(12)

        # === 1. Capacity for residence/leave ===
        details_outer_layout.addWidget(QLabel("<b>Capacity for residence/leave:</b>"))
        self.capacity_group = QButtonGroup(self)
        capacity_row = QHBoxLayout()
        self.capacity_yes = QRadioButton("Yes")
        self.capacity_no = QRadioButton("No")
        self.capacity_group.addButton(self.capacity_yes)
        self.capacity_group.addButton(self.capacity_no)
        capacity_row.addWidget(self.capacity_yes)
        capacity_row.addWidget(self.capacity_no)
        capacity_row.addStretch()
        details_outer_layout.addLayout(capacity_row)

        # === Container for No path (DoLS) ===
        self.dols_container = QFrame()
        self.dols_container.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #f59e0b; border-radius: 6px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        dols_layout = QVBoxLayout(self.dols_container)
        dols_layout.setContentsMargins(10, 8, 10, 8)
        dols_layout.setSpacing(6)

        dols_layout.addWidget(QLabel("<b>Plan for DoLS on discharge:</b>"))
        self.dols_group = QButtonGroup(self)
        dols_row = QHBoxLayout()
        self.dols_yes = QRadioButton("Yes")
        self.dols_no = QRadioButton("No")
        self.dols_group.addButton(self.dols_yes)
        self.dols_group.addButton(self.dols_no)
        self.dols_yes.toggled.connect(self._update_escorted_overnight_preview)
        self.dols_no.toggled.connect(self._update_escorted_overnight_preview)
        dols_row.addWidget(self.dols_yes)
        dols_row.addWidget(self.dols_no)
        dols_row.addStretch()
        dols_layout.addLayout(dols_row)

        self.dols_container.setVisible(False)
        details_outer_layout.addWidget(self.dols_container)

        # === Container for Yes path (Initial testing) ===
        self.initial_testing_container = QFrame()
        self.initial_testing_container.setStyleSheet("QFrame { background: #d1fae5; border: 1px solid #10b981; border-radius: 6px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        testing_layout = QVBoxLayout(self.initial_testing_container)
        testing_layout.setContentsMargins(10, 8, 10, 8)
        testing_layout.setSpacing(6)

        testing_layout.addWidget(QLabel("<b>Is this initial testing before further extended leave?</b>"))
        self.initial_testing_group = QButtonGroup(self)
        testing_row = QHBoxLayout()
        self.initial_testing_yes = QRadioButton("Yes")
        self.initial_testing_no = QRadioButton("No")
        self.initial_testing_group.addButton(self.initial_testing_yes)
        self.initial_testing_group.addButton(self.initial_testing_no)
        self.initial_testing_yes.toggled.connect(self._update_escorted_overnight_preview)
        self.initial_testing_no.toggled.connect(self._update_escorted_overnight_preview)
        testing_row.addWidget(self.initial_testing_yes)
        testing_row.addWidget(self.initial_testing_no)
        testing_row.addStretch()
        testing_layout.addLayout(testing_row)

        self.initial_testing_container.setVisible(False)
        details_outer_layout.addWidget(self.initial_testing_container)

        # Connect capacity radios to show/hide containers
        def on_capacity_changed():
            if self.capacity_no.isChecked():
                self.dols_container.setVisible(True)
                self.initial_testing_container.setVisible(False)
            elif self.capacity_yes.isChecked():
                self.dols_container.setVisible(False)
                self.initial_testing_container.setVisible(True)
            else:
                self.dols_container.setVisible(False)
                self.initial_testing_container.setVisible(False)
            self._update_escorted_overnight_preview()

        self.capacity_yes.toggled.connect(on_capacity_changed)
        self.capacity_no.toggled.connect(on_capacity_changed)

        # === Discharge Plan Section (always visible within details) ===
        discharge_frame = QFrame()
        discharge_frame.setStyleSheet("QFrame { background: #e0e7ff; border: 1px solid #6366f1; border-radius: 6px; } QLabel { background: transparent; border: none; } QCheckBox { background: transparent; border: none; }")
        discharge_layout = QVBoxLayout(discharge_frame)
        discharge_layout.setContentsMargins(10, 8, 10, 8)
        discharge_layout.setSpacing(6)

        discharge_layout.addWidget(QLabel("<b>How does this fit into discharge plan?</b>"))
        self.discharge_dols_cb = QCheckBox("Discharge on DoLS")
        self.discharge_unescorted_cb = QCheckBox("Discharge after unescorted leave")
        self.discharge_initial_cb = QCheckBox("Initial testing for overnight leave")

        for cb in [self.discharge_dols_cb, self.discharge_unescorted_cb, self.discharge_initial_cb]:
            cb.stateChanged.connect(self._update_escorted_overnight_preview)
            discharge_layout.addWidget(cb)

        details_outer_layout.addWidget(discharge_frame)

        # Hide details container initially
        self.escorted_overnight_details_container.setVisible(False)
        input_layout.addWidget(self.escorted_overnight_details_container)

        # Connect Yes/NA to show/hide details
        def on_escorted_overnight_main_changed():
            if self.escorted_overnight_yes.isChecked():
                self.escorted_overnight_details_container.setVisible(True)
            else:
                self.escorted_overnight_details_container.setVisible(False)
            self._update_escorted_overnight_preview()

        self.escorted_overnight_yes.toggled.connect(on_escorted_overnight_main_changed)
        self.escorted_overnight_na.toggled.connect(on_escorted_overnight_main_changed)

        layout.addWidget(input_frame)

        layout.addStretch()
        popup.setWidget(content)
        return popup

    def _update_escorted_overnight_preview(self):
        """Update the escorted overnight preview with generated narrative."""
        if not hasattr(self, 'escorted_overnight_preview'):
            return

        # Check if N/A selected
        if hasattr(self, 'escorted_overnight_na') and self.escorted_overnight_na.isChecked():
            self.escorted_overnight_preview.setPlainText("Escorted overnight leave is not applicable.")
            return

        if not (hasattr(self, 'escorted_overnight_yes') and self.escorted_overnight_yes.isChecked()):
            self.escorted_overnight_preview.setPlainText("")
            return

        p = self._get_pronouns()
        parts = []

        # Capacity section
        if hasattr(self, 'capacity_yes') and self.capacity_yes.isChecked():
            parts.append(f"{p['subj']} {p['have']} capacity to make decisions about {p['pos']} residence and leave.")
            # Initial testing follow-up
            if hasattr(self, 'initial_testing_yes') and self.initial_testing_yes.isChecked():
                parts.append(f"This escorted overnight leave represents initial testing before further extended leave is considered.")
            elif hasattr(self, 'initial_testing_no') and self.initial_testing_no.isChecked():
                parts.append(f"{p['subj']} {p['have']} already completed initial testing and this leave continues {p['pos']} rehabilitation pathway.")
        elif hasattr(self, 'capacity_no') and self.capacity_no.isChecked():
            parts.append(f"{p['subj']} {p['do']} not have capacity to make decisions about {p['pos']} residence and leave.")
            # DoLS follow-up
            if hasattr(self, 'dols_yes') and self.dols_yes.isChecked():
                parts.append(f"There is a plan for Deprivation of Liberty Safeguards (DoLS) to be in place on discharge.")
            elif hasattr(self, 'dols_no') and self.dols_no.isChecked():
                parts.append(f"DoLS will not be required on discharge as alternative arrangements are in place.")

        # Discharge plan section
        discharge_parts = []
        if hasattr(self, 'discharge_dols_cb') and self.discharge_dols_cb.isChecked():
            discharge_parts.append("discharge under DoLS")
        if hasattr(self, 'discharge_unescorted_cb') and self.discharge_unescorted_cb.isChecked():
            discharge_parts.append("discharge following a period of unescorted leave")
        if hasattr(self, 'discharge_initial_cb') and self.discharge_initial_cb.isChecked():
            discharge_parts.append("initial testing for overnight leave arrangements")

        if discharge_parts:
            if len(discharge_parts) == 1:
                parts.append(f"This leave fits into {p['pos']} discharge plan as part of {discharge_parts[0]}.")
            elif len(discharge_parts) == 2:
                parts.append(f"This leave fits into {p['pos']} discharge plan as part of {discharge_parts[0]} and {discharge_parts[1]}.")
            else:
                parts.append(f"This leave fits into {p['pos']} discharge plan as part of {', '.join(discharge_parts[:-1])}, and {discharge_parts[-1]}.")

        self.escorted_overnight_preview.setPlainText(" ".join(parts))

    def _create_compassionate_popup(self) -> QWidget:
        """Create compassionate leave popup with structured questions."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { border: none; background: white; } QWidget { background: white; }")

        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        # === PREVIEW SECTION (fixed height) ===
        preview_container = QFrame()
        preview_container.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; }")
        preview_container.setMaximumHeight(160)
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 8, 12, 8)
        preview_layout.setSpacing(4)

        preview_header_row = QHBoxLayout()
        preview_header_row.setContentsMargins(0, 0, 0, 0)

        preview_header = QLabel("Preview")
        preview_header.setStyleSheet("font-size: 13px; font-weight: 600; color: #991b1b;")
        preview_header_row.addWidget(preview_header)
        preview_header_row.addStretch()

        self.compassionate_send_btn = QPushButton("Send to Card")
        self.compassionate_send_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 5px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        self.compassionate_send_btn.clicked.connect(lambda: self._send_to_card("compassionate"))
        preview_header_row.addWidget(self.compassionate_send_btn)

        preview_layout.addLayout(preview_header_row)

        self.compassionate_preview = QTextEdit()
        self.compassionate_preview.setReadOnly(True)
        self.compassionate_preview.setMinimumHeight(60)
        self.compassionate_preview.setMaximumHeight(100)
        self.compassionate_preview.setStyleSheet("""
            QTextEdit { background: #1f2937; border: none; border-radius: 4px; padding: 8px; font-size: 12px; color: white; }
        """)
        preview_layout.addWidget(self.compassionate_preview)

        layout.addWidget(preview_container)

        # === INPUT SECTION ===
        input_frame = QFrame()
        input_frame.setStyleSheet("QFrame { background: #f3f4f6; border-radius: 8px; } QLabel { background: #f3f4f6; border: none; } QRadioButton, QCheckBox { background: #f3f4f6; border: none; }")
        input_layout = QVBoxLayout(input_frame)
        input_layout.setContentsMargins(12, 12, 12, 12)
        input_layout.setSpacing(12)

        # === Compassionate Leave Required: Yes/NA ===
        input_layout.addWidget(QLabel("<b>Compassionate Leave Required:</b>"))
        self.compassionate_main_group = QButtonGroup(self)
        main_row = QHBoxLayout()
        self.compassionate_yes = QRadioButton("Yes")
        self.compassionate_na = QRadioButton("N/A")
        self.compassionate_main_group.addButton(self.compassionate_yes)
        self.compassionate_main_group.addButton(self.compassionate_na)
        main_row.addWidget(self.compassionate_yes)
        main_row.addWidget(self.compassionate_na)
        main_row.addStretch()
        input_layout.addLayout(main_row)

        # === Details container (shown when Yes) ===
        self.compassionate_details_container = QFrame()
        self.compassionate_details_container.setStyleSheet("QFrame { background: transparent; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        details_layout = QVBoxLayout(self.compassionate_details_container)
        details_layout.setContentsMargins(0, 0, 0, 0)
        details_layout.setSpacing(10)

        # Can this be virtual?
        virtual_frame = QFrame()
        virtual_frame.setStyleSheet("QFrame { background: #e0e7ff; border: 1px solid #6366f1; border-radius: 6px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        virtual_layout = QVBoxLayout(virtual_frame)
        virtual_layout.setContentsMargins(10, 8, 10, 8)
        virtual_layout.setSpacing(6)

        virtual_layout.addWidget(QLabel("<b>Can this be virtual?</b>"))
        self.virtual_group = QButtonGroup(self)
        virtual_row = QHBoxLayout()
        self.virtual_yes = QRadioButton("Yes")
        self.virtual_no = QRadioButton("No")
        self.virtual_group.addButton(self.virtual_yes)
        self.virtual_group.addButton(self.virtual_no)
        self.virtual_yes.toggled.connect(self._update_compassionate_preview)
        self.virtual_no.toggled.connect(self._update_compassionate_preview)
        virtual_row.addWidget(self.virtual_yes)
        virtual_row.addWidget(self.virtual_no)
        virtual_row.addStretch()
        virtual_layout.addLayout(virtual_row)

        # Reasons (shown when virtual=No)
        self.virtual_reasons_container = QFrame()
        self.virtual_reasons_container.setStyleSheet("QFrame { background: transparent; }")
        reasons_layout = QVBoxLayout(self.virtual_reasons_container)
        reasons_layout.setContentsMargins(0, 4, 0, 0)
        reasons_layout.setSpacing(4)
        reasons_layout.addWidget(QLabel("Reasons why virtual is not possible:"))
        self.virtual_reasons = QTextEdit()
        self.virtual_reasons.setPlaceholderText("Explain why virtual attendance is not possible...")
        self.virtual_reasons.setMaximumHeight(80)
        self.virtual_reasons.setStyleSheet("QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 4px; padding: 6px; font-size: 11px; }")
        self.virtual_reasons.textChanged.connect(self._update_compassionate_preview)
        reasons_layout.addWidget(self.virtual_reasons)
        self.virtual_reasons_container.setVisible(False)
        virtual_layout.addWidget(self.virtual_reasons_container)

        # Connect virtual radios
        def on_virtual_changed():
            self.virtual_reasons_container.setVisible(self.virtual_no.isChecked())
            self._update_compassionate_preview()

        self.virtual_yes.toggled.connect(on_virtual_changed)
        self.virtual_no.toggled.connect(on_virtual_changed)

        details_layout.addWidget(virtual_frame)

        # Details of people affected
        people_frame = QFrame()
        people_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #f59e0b; border-radius: 6px; } QLabel { background: transparent; border: none; }")
        people_layout = QVBoxLayout(people_frame)
        people_layout.setContentsMargins(10, 8, 10, 8)
        people_layout.setSpacing(6)

        people_layout.addWidget(QLabel("<b>Details of people affected:</b>"))
        self.people_affected = QTextEdit()
        self.people_affected.setPlaceholderText("Enter details of family members or others affected...")
        self.people_affected.setMaximumHeight(80)
        self.people_affected.setStyleSheet("QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 4px; padding: 6px; font-size: 11px; }")
        self.people_affected.textChanged.connect(self._update_compassionate_preview)
        people_layout.addWidget(self.people_affected)

        details_layout.addWidget(people_frame)

        # Urgent?
        urgent_frame = QFrame()
        urgent_frame.setStyleSheet("QFrame { background: #fee2e2; border: 1px solid #ef4444; border-radius: 6px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        urgent_layout = QVBoxLayout(urgent_frame)
        urgent_layout.setContentsMargins(10, 8, 10, 8)
        urgent_layout.setSpacing(6)

        urgent_layout.addWidget(QLabel("<b>Urgent:</b>"))
        self.urgent_group = QButtonGroup(self)
        urgent_row = QHBoxLayout()
        self.urgent_yes = QRadioButton("Yes")
        self.urgent_no = QRadioButton("No")
        self.urgent_group.addButton(self.urgent_yes)
        self.urgent_group.addButton(self.urgent_no)
        self.urgent_yes.toggled.connect(self._update_compassionate_preview)
        self.urgent_no.toggled.connect(self._update_compassionate_preview)
        urgent_row.addWidget(self.urgent_yes)
        urgent_row.addWidget(self.urgent_no)
        urgent_row.addStretch()
        urgent_layout.addLayout(urgent_row)

        details_layout.addWidget(urgent_frame)

        self.compassionate_details_container.setVisible(False)
        input_layout.addWidget(self.compassionate_details_container)

        # Connect main Yes/NA
        def on_compassionate_main_changed():
            self.compassionate_details_container.setVisible(self.compassionate_yes.isChecked())
            self._update_compassionate_preview()

        self.compassionate_yes.toggled.connect(on_compassionate_main_changed)
        self.compassionate_na.toggled.connect(on_compassionate_main_changed)

        layout.addWidget(input_frame)

        layout.addStretch()
        popup.setWidget(content)
        return popup

    def _update_compassionate_preview(self):
        """Update the compassionate leave preview with generated narrative."""
        if not hasattr(self, 'compassionate_preview'):
            return

        p = self._get_pronouns()
        parts = []

        if hasattr(self, 'compassionate_na') and self.compassionate_na.isChecked():
            parts.append("Compassionate leave is not applicable at this time.")
        elif hasattr(self, 'compassionate_yes') and self.compassionate_yes.isChecked():
            parts.append("Compassionate leave is being requested.")

            # Virtual
            if hasattr(self, 'virtual_yes') and self.virtual_yes.isChecked():
                parts.append("This visit can be conducted virtually.")
            elif hasattr(self, 'virtual_no') and self.virtual_no.isChecked():
                reasons = self.virtual_reasons.toPlainText().strip() if hasattr(self, 'virtual_reasons') else ""
                if reasons:
                    parts.append(f"A virtual visit is not possible because {reasons.lower() if not reasons[0].isupper() else reasons}")
                else:
                    parts.append("A virtual visit is not possible in this case.")

            # People affected
            if hasattr(self, 'people_affected'):
                people = self.people_affected.toPlainText().strip()
                if people:
                    parts.append(f"The people affected are: {people}")

            # Urgent
            if hasattr(self, 'urgent_yes') and self.urgent_yes.isChecked():
                parts.append("This request is urgent and requires expedited consideration.")
            elif hasattr(self, 'urgent_no') and self.urgent_no.isChecked():
                parts.append("This request is not urgent.")

        self.compassionate_preview.setPlainText(" ".join(parts))

    def _create_leave_report_popup(self) -> QWidget:
        """Create leave report popup - 3g. Leave report last 2 years."""
        # Initialize state storage for escorted/unescorted entries
        self._leave_escorted_state = {}
        self._leave_unescorted_state = {}
        self._leave_switching = False  # Flag to prevent recursive updates

        popup = QWidget()
        main_layout = QVBoxLayout(popup)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(8)

        # === HEADER ===
        header_label = QLabel("<b>3g. Leave report last 2 years</b>")
        header_label.setStyleSheet("font-size: 14px; color: #1f2937; padding: 8px; border: none;")
        main_layout.addWidget(header_label)

        # === FIXED TOP: Preview + Send to Card ===
        top_container = QFrame()
        top_container.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; }")
        top_container.setMaximumHeight(160)
        top_layout = QVBoxLayout(top_container)
        top_layout.setContentsMargins(12, 8, 12, 8)
        top_layout.setSpacing(4)

        # Header row with Preview label and Send to Card button
        header_row = QHBoxLayout()
        preview_header = QLabel("Preview")
        preview_header.setStyleSheet("font-size: 13px; font-weight: 600; color: #991b1b; border: none;")
        header_row.addWidget(preview_header)
        header_row.addStretch()

        send_btn = QPushButton("Send to Card")
        send_btn.setStyleSheet("QPushButton { background: #991b1b; color: white; padding: 8px 16px; border: none; border-radius: 6px; font-weight: 600; } QPushButton:hover { background: #7f1d1d; }")
        send_btn.clicked.connect(lambda: self._send_to_card("leave_report"))
        header_row.addWidget(send_btn)
        top_layout.addLayout(header_row)

        self.leave_report_preview = QTextEdit()
        self.leave_report_preview.setReadOnly(True)
        self.leave_report_preview.setMinimumHeight(80)
        self.leave_report_preview.setMaximumHeight(120)
        self.leave_report_preview.setStyleSheet("QTextEdit { background: #1f2937; border: none; border-radius: 4px; padding: 8px; font-size: 12px; color: white; }")
        top_layout.addWidget(self.leave_report_preview)

        main_layout.addWidget(top_container)

        # === SCROLLABLE CONTENT ===
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setStyleSheet("QScrollArea { border: none; background: white; }")

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(8)

        # === INPUT FIELDS SECTION ===
        self.leave_input_section = CollapsibleSection("Input Fields", start_collapsed=False)
        self.leave_input_section.set_content_height(380)

        input_content = QFrame()
        input_content.setStyleSheet("QFrame { background: #f3f4f6; } QLabel { background: transparent; border: none; } QRadioButton, QCheckBox { background: transparent; border: none; }")
        input_layout = QVBoxLayout(input_content)
        input_layout.setContentsMargins(10, 10, 10, 10)
        input_layout.setSpacing(8)

        # === ESCORTED/UNESCORTED TOGGLE ===
        escort_frame = QFrame()
        escort_frame.setStyleSheet("QFrame { background: #dbeafe; border: 2px solid #3b82f6; border-radius: 8px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        escort_layout = QHBoxLayout(escort_frame)
        escort_layout.setContentsMargins(12, 8, 12, 8)
        escort_layout.setSpacing(16)

        escort_label = QLabel("<b>Leave Type:</b>")
        escort_label.setStyleSheet("font-size: 13px; color: #1e40af;")
        escort_layout.addWidget(escort_label)

        self.leave_escort_group = QButtonGroup(self)
        self.leave_escorted_radio = QRadioButton("Escorted")
        self.leave_unescorted_radio = QRadioButton("Unescorted")
        self.leave_escorted_radio.setStyleSheet("font-weight: 600; font-size: 13px; color: #1e40af;")
        self.leave_unescorted_radio.setStyleSheet("font-weight: 600; font-size: 13px; color: #1e40af;")
        self.leave_escort_group.addButton(self.leave_escorted_radio)
        self.leave_escort_group.addButton(self.leave_unescorted_radio)
        self.leave_escorted_radio.setChecked(True)  # Default to escorted
        escort_layout.addWidget(self.leave_escorted_radio)
        escort_layout.addWidget(self.leave_unescorted_radio)
        escort_layout.addStretch()

        input_layout.addWidget(escort_frame)

        # Row 1: Leaves, Frequency, Duration
        dropdown_style = """
            QComboBox {
                background: white;
                color: black;
                padding: 6px 10px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                font-size: 12px;
            }
            QComboBox::drop-down { width: 20px; }
            QComboBox QAbstractItemView { background: white; color: black; selection-background-color: #dbeafe; }
        """

        row1 = QHBoxLayout()
        row1.addWidget(QLabel("Leaves:"))
        self.leaves_dropdown = QComboBox()
        self.leaves_dropdown.addItems([str(i) for i in range(1, 8)])
        self.leaves_dropdown.setMinimumWidth(60)
        self.leaves_dropdown.setStyleSheet(dropdown_style)
        self.leaves_dropdown.currentIndexChanged.connect(self._update_leave_report_preview)
        row1.addWidget(self.leaves_dropdown)
        row1.addWidget(QLabel("Freq:"))
        self.frequency_dropdown = QComboBox()
        self.frequency_dropdown.addItems(["Weekly", "2 weekly", "3 weekly", "Monthly", "2 monthly"])
        self.frequency_dropdown.setMinimumWidth(100)
        self.frequency_dropdown.setStyleSheet(dropdown_style)
        self.frequency_dropdown.currentIndexChanged.connect(self._update_leave_report_preview)
        row1.addWidget(self.frequency_dropdown)
        row1.addWidget(QLabel("Dur:"))
        self.duration_dropdown = QComboBox()
        self.duration_dropdown.addItems(["30 mins", "1 hour", "2 hours", "3 hours", "4 hours", "5 hours", "6 hours", "7 hours", "8 hours"])
        self.duration_dropdown.setMinimumWidth(90)
        self.duration_dropdown.setStyleSheet(dropdown_style)
        self.duration_dropdown.currentIndexChanged.connect(self._update_leave_report_preview)
        row1.addWidget(self.duration_dropdown)
        row1.addStretch()
        input_layout.addLayout(row1)

        # Leave types with individual linked sliders
        weight_frame = QFrame()
        weight_frame.setStyleSheet("QFrame { background: #e0e7ff; border: 1px solid #6366f1; border-radius: 6px; } QLabel { background: transparent; border: none; } QCheckBox { background: transparent; border: none; }")
        weight_layout = QVBoxLayout(weight_frame)
        weight_layout.setContentsMargins(10, 8, 10, 8)
        weight_layout.setSpacing(6)

        weight_layout.addWidget(QLabel("<b>Leave types & weighting:</b>"))

        # Store leave type data: checkbox, slider, label
        self.leave_type_widgets = {}
        self._updating_sliders = False

        leave_types = [
            ("ground", "Ground"),
            ("local", "Local community"),
            ("community", "Community"),
            ("extended", "Extended community"),
            ("overnight", "Overnight")
        ]

        for key, label in leave_types:
            row = QHBoxLayout()
            row.setSpacing(8)

            cb = QCheckBox(label)
            cb.setFixedWidth(130)
            row.addWidget(cb)

            slider = NoWheelSlider(Qt.Horizontal)
            slider.setMinimum(0)
            slider.setMaximum(100)
            slider.setValue(0)
            slider.setEnabled(False)
            slider.setStyleSheet("QSlider::groove:horizontal { background: #c7d2fe; height: 6px; border-radius: 3px; } QSlider::handle:horizontal { background: #4f46e5; width: 14px; margin: -4px 0; border-radius: 7px; } QSlider:disabled::groove:horizontal { background: #e5e7eb; } QSlider:disabled::handle:horizontal { background: #9ca3af; }")
            row.addWidget(slider, 1)

            pct_label = QLabel("0%")
            pct_label.setFixedWidth(35)
            pct_label.setStyleSheet("font-weight: 600; color: #4f46e5;")
            row.addWidget(pct_label)

            self.leave_type_widgets[key] = {"cb": cb, "slider": slider, "label": pct_label}

            cb.stateChanged.connect(lambda state, k=key: self._on_leave_type_toggled(k, state))
            slider.valueChanged.connect(lambda val, k=key: self._on_leave_slider_changed(k, val))

            weight_layout.addLayout(row)

        input_layout.addWidget(weight_frame)

        # Other leave types
        other_frame = QFrame()
        other_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #f59e0b; border-radius: 6px; } QLabel { background: transparent; border: none; } QCheckBox { background: transparent; border: none; }")
        other_layout = QVBoxLayout(other_frame)
        other_layout.setContentsMargins(10, 8, 10, 8)
        other_layout.setSpacing(6)

        other_layout.addWidget(QLabel("<b>Other leave:</b>"))
        other_row = QHBoxLayout()
        self.leave_medical_cb = QCheckBox("Medical")
        self.leave_court_cb = QCheckBox("Court")
        self.leave_compassionate_cb = QCheckBox("Compassionate")
        for cb in [self.leave_medical_cb, self.leave_court_cb, self.leave_compassionate_cb]:
            cb.stateChanged.connect(self._update_leave_report_preview)
            other_row.addWidget(cb)
        other_row.addStretch()
        other_layout.addLayout(other_row)

        input_layout.addWidget(other_frame)

        # Leave suspended
        suspended_frame = QFrame()
        suspended_frame.setStyleSheet("QFrame { background: #fee2e2; border: 1px solid #ef4444; border-radius: 6px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        suspended_layout = QVBoxLayout(suspended_frame)
        suspended_layout.setContentsMargins(10, 8, 10, 8)
        suspended_layout.setSpacing(6)

        suspended_layout.addWidget(QLabel("<b>Leave ever suspended:</b>"))
        self.suspended_group = QButtonGroup(self)
        suspended_row = QHBoxLayout()
        self.suspended_yes = QRadioButton("Yes")
        self.suspended_no = QRadioButton("No")
        self.suspended_group.addButton(self.suspended_yes)
        self.suspended_group.addButton(self.suspended_no)
        suspended_row.addWidget(self.suspended_yes)
        suspended_row.addWidget(self.suspended_no)
        suspended_row.addStretch()
        suspended_layout.addLayout(suspended_row)

        # Suspension details (shown when Yes) - taller
        self.suspension_details_container = QFrame()
        self.suspension_details_container.setStyleSheet("QFrame { background: transparent; } QLabel { background: transparent; border: none; }")
        susp_details_layout = QVBoxLayout(self.suspension_details_container)
        susp_details_layout.setContentsMargins(0, 4, 0, 0)
        susp_details_layout.setSpacing(4)
        susp_details_layout.addWidget(QLabel("Details:"))
        self.suspension_details = QTextEdit()
        self.suspension_details.setPlaceholderText("Describe the circumstances of leave suspension...")
        self.suspension_details.setMinimumHeight(80)
        self.suspension_details.setStyleSheet("QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 4px; padding: 6px; font-size: 11px; }")
        self.suspension_details.textChanged.connect(self._update_leave_report_preview)
        susp_details_layout.addWidget(self.suspension_details)
        self.suspension_details_container.setVisible(False)
        suspended_layout.addWidget(self.suspension_details_container)

        def on_suspended_changed():
            self.suspension_details_container.setVisible(self.suspended_yes.isChecked())
            self._update_leave_report_preview()

        self.suspended_yes.toggled.connect(on_suspended_changed)
        self.suspended_no.toggled.connect(on_suspended_changed)

        input_layout.addWidget(suspended_frame)

        # Connect escorted/unescorted toggle to save/restore state
        def on_escort_type_changed():
            if self._leave_switching:
                return
            self._leave_switching = True
            try:
                if self.leave_escorted_radio.isChecked():
                    # Save unescorted state, restore escorted state
                    self._save_leave_state(self._leave_unescorted_state)
                    self._restore_leave_state(self._leave_escorted_state)
                else:
                    # Save escorted state, restore unescorted state
                    self._save_leave_state(self._leave_escorted_state)
                    self._restore_leave_state(self._leave_unescorted_state)
                self._update_leave_report_preview()
            finally:
                self._leave_switching = False

        self.leave_escorted_radio.toggled.connect(on_escort_type_changed)

        self.leave_input_section.set_content(input_content)
        scroll_layout.addWidget(self.leave_input_section)

        # === IMPORTED NOTES SECTION (like ASR section 6 - light orange) ===
        self.leave_import_section = CollapsibleSection("Imported Notes", start_collapsed=True)
        self.leave_import_section.set_content_height(350)

        import_panel = QFrame()
        import_panel.setStyleSheet("QFrame { background: #fffbeb; border: 1px solid #fcd34d; border-radius: 8px; }")
        import_panel_layout = QVBoxLayout(import_panel)
        import_panel_layout.setContentsMargins(10, 10, 10, 10)
        import_panel_layout.setSpacing(8)

        # Subtitle with count
        self.leave_import_subtitle = QLabel("Leave evidence from notes (last 2 years)")
        self.leave_import_subtitle.setStyleSheet("font-size: 11px; font-weight: 600; color: #92400e; background: transparent;")
        import_panel_layout.addWidget(self.leave_import_subtitle)

        # Filter tags row (populated dynamically)
        self.leave_filter_row = QHBoxLayout()
        self.leave_filter_row.setSpacing(6)
        self.leave_filter_row.addStretch()
        import_panel_layout.addLayout(self.leave_filter_row)

        # Scrollable content area
        leave_import_scroll = QScrollArea()
        leave_import_scroll.setWidgetResizable(True)
        leave_import_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        leave_import_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        leave_import_scroll.setStyleSheet("QScrollArea { background: #fffbeb; border: none; }")
        leave_import_scroll.setMinimumHeight(180)

        # Content widget for entries
        self.leave_import_content = QWidget()
        self.leave_import_content.setStyleSheet("background: transparent;")
        self.leave_import_layout = QVBoxLayout(self.leave_import_content)
        self.leave_import_layout.setContentsMargins(4, 4, 4, 4)
        self.leave_import_layout.setSpacing(6)
        leave_import_scroll.setWidget(self.leave_import_content)

        # Placeholder
        self.leave_import_placeholder = QLabel("No entries found. Use Import Data to load clinical notes.")
        self.leave_import_placeholder.setStyleSheet("color: #92400e; font-style: italic; font-size: 11px; background: transparent;")
        self.leave_import_layout.addWidget(self.leave_import_placeholder)

        import_panel_layout.addWidget(leave_import_scroll, 1)

        # Button row
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(8)

        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet("""
            QPushButton { background: #fef3c7; color: #92400e; border: 1px solid #fcd34d; padding: 6px 12px; border-radius: 6px; font-weight: 500; font-size: 11px; }
            QPushButton:hover { background: #fde68a; }
        """)
        clear_btn.clicked.connect(self._clear_leave_imports)
        btn_layout.addWidget(clear_btn)

        btn_layout.addStretch()
        import_panel_layout.addLayout(btn_layout)

        # Storage for imported entries and checkboxes
        self.leave_imported_entries = []
        self.leave_import_checkboxes = []
        self._leave_current_filter = None
        self._leave_all_categorized = []

        self.leave_import_section.set_content(import_panel)
        scroll_layout.addWidget(self.leave_import_section)

        scroll_layout.addStretch()
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area, 1)

        return popup

    def _on_leave_type_toggled(self, key, state):
        """Handle leave type checkbox toggle - enable/disable slider and redistribute weights."""
        if not hasattr(self, 'leave_type_widgets'):
            return

        widgets = self.leave_type_widgets[key]
        is_checked = state == Qt.CheckState.Checked.value

        widgets["slider"].setEnabled(is_checked)

        if is_checked:
            self._add_leave_type_weight(key)
        else:
            self._remove_leave_type_weight(key)

        self._update_leave_report_preview()

    def _add_leave_type_weight(self, new_key):
        """Add a new leave type, taking weight proportionally from existing checked items."""
        self._updating_sliders = True

        other_checked = [(k, w["slider"].value()) for k, w in self.leave_type_widgets.items()
                         if w["cb"].isChecked() and k != new_key]

        if not other_checked:
            slider = self.leave_type_widgets[new_key]["slider"]
            slider.blockSignals(True)
            slider.setValue(100)
            slider.blockSignals(False)
            self.leave_type_widgets[new_key]["label"].setText("100%")
        else:
            new_item_share = 1
            remaining = 100 - new_item_share

            total_existing = sum(v for _, v in other_checked)
            if total_existing == 0:
                total_existing = 100

            for k, old_val in other_checked:
                if total_existing > 0:
                    new_val = int((old_val / total_existing) * remaining)
                else:
                    new_val = remaining // len(other_checked)
                slider = self.leave_type_widgets[k]["slider"]
                slider.blockSignals(True)
                slider.setValue(new_val)
                slider.blockSignals(False)
                self.leave_type_widgets[k]["label"].setText(f"{new_val}%")

            slider = self.leave_type_widgets[new_key]["slider"]
            slider.blockSignals(True)
            slider.setValue(new_item_share)
            slider.blockSignals(False)
            self.leave_type_widgets[new_key]["label"].setText(f"{new_item_share}%")

        self._updating_sliders = False

    def _remove_leave_type_weight(self, removed_key):
        """Remove a leave type, distributing its weight to remaining checked items."""
        self._updating_sliders = True

        removed_weight = self.leave_type_widgets[removed_key]["slider"].value()

        slider = self.leave_type_widgets[removed_key]["slider"]
        slider.blockSignals(True)
        slider.setValue(0)
        slider.blockSignals(False)
        self.leave_type_widgets[removed_key]["label"].setText("0%")

        remaining_checked = [(k, w["slider"].value()) for k, w in self.leave_type_widgets.items()
                            if w["cb"].isChecked() and k != removed_key]

        if remaining_checked:
            total_remaining = sum(v for _, v in remaining_checked)
            if total_remaining == 0:
                per_item = removed_weight // len(remaining_checked)
                extra = removed_weight % len(remaining_checked)
                for i, (k, _) in enumerate(remaining_checked):
                    new_val = per_item + (1 if i < extra else 0)
                    slider = self.leave_type_widgets[k]["slider"]
                    slider.blockSignals(True)
                    slider.setValue(new_val)
                    slider.blockSignals(False)
                    self.leave_type_widgets[k]["label"].setText(f"{new_val}%")
            else:
                for k, old_val in remaining_checked:
                    proportion = old_val / total_remaining
                    new_val = old_val + int(removed_weight * proportion)
                    slider = self.leave_type_widgets[k]["slider"]
                    slider.blockSignals(True)
                    slider.setValue(min(100, new_val))
                    slider.blockSignals(False)
                    self.leave_type_widgets[k]["label"].setText(f"{min(100, new_val)}%")

        self._updating_sliders = False

    def _on_leave_slider_changed(self, key, value):
        """Handle slider change - redistribute remaining weight proportionally among other checked items."""
        if self._updating_sliders or not hasattr(self, 'leave_type_widgets'):
            return

        self._updating_sliders = True

        self.leave_type_widgets[key]["label"].setText(f"{value}%")

        others = [(k, w["slider"].value()) for k, w in self.leave_type_widgets.items()
                  if w["cb"].isChecked() and k != key]

        if others:
            remaining = 100 - value
            total_others = sum(v for _, v in others)

            if total_others > 0:
                for k, old_val in others:
                    proportion = old_val / total_others
                    new_val = int(remaining * proportion)
                    slider = self.leave_type_widgets[k]["slider"]
                    slider.blockSignals(True)
                    slider.setValue(new_val)
                    slider.blockSignals(False)
                    self.leave_type_widgets[k]["label"].setText(f"{new_val}%")
            else:
                per_item = remaining // len(others)
                extra = remaining % len(others)
                for i, (k, _) in enumerate(others):
                    new_val = per_item + (1 if i < extra else 0)
                    slider = self.leave_type_widgets[k]["slider"]
                    slider.blockSignals(True)
                    slider.setValue(new_val)
                    slider.blockSignals(False)
                    self.leave_type_widgets[k]["label"].setText(f"{new_val}%")

        self._updating_sliders = False
        self._update_leave_report_preview()

    def _save_leave_state(self, state_dict):
        """Save current leave report form state to a dictionary."""
        state_dict['leaves'] = self.leaves_dropdown.currentIndex() if hasattr(self, 'leaves_dropdown') else 0
        state_dict['frequency'] = self.frequency_dropdown.currentIndex() if hasattr(self, 'frequency_dropdown') else 0
        state_dict['duration'] = self.duration_dropdown.currentIndex() if hasattr(self, 'duration_dropdown') else 0
        state_dict['leave_types'] = {}
        if hasattr(self, 'leave_type_widgets'):
            for key, widgets in self.leave_type_widgets.items():
                state_dict['leave_types'][key] = {
                    'checked': widgets['cb'].isChecked(),
                    'value': widgets['slider'].value()
                }
        state_dict['medical'] = self.leave_medical_cb.isChecked() if hasattr(self, 'leave_medical_cb') else False
        state_dict['court'] = self.leave_court_cb.isChecked() if hasattr(self, 'leave_court_cb') else False
        state_dict['compassionate'] = self.leave_compassionate_cb.isChecked() if hasattr(self, 'leave_compassionate_cb') else False
        state_dict['suspended_yes'] = self.suspended_yes.isChecked() if hasattr(self, 'suspended_yes') else False
        state_dict['suspended_no'] = self.suspended_no.isChecked() if hasattr(self, 'suspended_no') else False
        state_dict['suspension_details'] = self.suspension_details.toPlainText() if hasattr(self, 'suspension_details') else ""

    def _restore_leave_state(self, state_dict):
        """Restore leave report form state from a dictionary."""
        if not state_dict:
            # Reset to defaults
            if hasattr(self, 'leaves_dropdown'):
                self.leaves_dropdown.setCurrentIndex(0)
            if hasattr(self, 'frequency_dropdown'):
                self.frequency_dropdown.setCurrentIndex(0)
            if hasattr(self, 'duration_dropdown'):
                self.duration_dropdown.setCurrentIndex(0)
            if hasattr(self, 'leave_type_widgets'):
                for widgets in self.leave_type_widgets.values():
                    widgets['cb'].setChecked(False)
                    widgets['slider'].setValue(0)
            if hasattr(self, 'leave_medical_cb'):
                self.leave_medical_cb.setChecked(False)
            if hasattr(self, 'leave_court_cb'):
                self.leave_court_cb.setChecked(False)
            if hasattr(self, 'leave_compassionate_cb'):
                self.leave_compassionate_cb.setChecked(False)
            if hasattr(self, 'suspended_yes'):
                self.suspended_yes.setChecked(False)
            if hasattr(self, 'suspended_no'):
                self.suspended_no.setChecked(False)
            if hasattr(self, 'suspension_details'):
                self.suspension_details.clear()
            return

        if hasattr(self, 'leaves_dropdown'):
            self.leaves_dropdown.setCurrentIndex(state_dict.get('leaves', 0))
        if hasattr(self, 'frequency_dropdown'):
            self.frequency_dropdown.setCurrentIndex(state_dict.get('frequency', 0))
        if hasattr(self, 'duration_dropdown'):
            self.duration_dropdown.setCurrentIndex(state_dict.get('duration', 0))
        if hasattr(self, 'leave_type_widgets'):
            leave_types = state_dict.get('leave_types', {})
            for key, widgets in self.leave_type_widgets.items():
                lt = leave_types.get(key, {})
                widgets['cb'].setChecked(lt.get('checked', False))
                widgets['slider'].setValue(lt.get('value', 0))
        if hasattr(self, 'leave_medical_cb'):
            self.leave_medical_cb.setChecked(state_dict.get('medical', False))
        if hasattr(self, 'leave_court_cb'):
            self.leave_court_cb.setChecked(state_dict.get('court', False))
        if hasattr(self, 'leave_compassionate_cb'):
            self.leave_compassionate_cb.setChecked(state_dict.get('compassionate', False))
        if state_dict.get('suspended_yes') and hasattr(self, 'suspended_yes'):
            self.suspended_yes.setChecked(True)
        elif state_dict.get('suspended_no') and hasattr(self, 'suspended_no'):
            self.suspended_no.setChecked(True)
        else:
            if hasattr(self, 'suspended_yes'):
                self.suspended_yes.setChecked(False)
            if hasattr(self, 'suspended_no'):
                self.suspended_no.setChecked(False)
        if hasattr(self, 'suspension_details'):
            self.suspension_details.setPlainText(state_dict.get('suspension_details', ''))

    def _generate_leave_report_from_state(self, state_dict, leave_type_label):
        """Generate leave report text from a state dictionary."""
        if not state_dict:
            return ""

        p = self._get_pronouns()
        parts = []

        # Get values from state
        leaves_idx = state_dict.get('leaves', 0)
        leaves = str(leaves_idx + 1)  # Index 0 = "1", etc.
        freq_options = ["weekly", "2 weekly", "3 weekly", "monthly", "2 monthly"]
        frequency = freq_options[state_dict.get('frequency', 0)]
        dur_options = ["30 mins", "1 hour", "2 hours", "3 hours", "4 hours", "5 hours", "6 hours", "7 hours", "8 hours"]
        duration = dur_options[state_dict.get('duration', 0)]

        leave_type_labels = {
            "ground": "ground leave",
            "local": "local community leave",
            "community": "community leave",
            "extended": "extended community leave",
            "overnight": "overnight leave"
        }

        # Get checked leave types with their weights
        leave_types = state_dict.get('leave_types', {})
        checked_types = []
        for key, lt in leave_types.items():
            if lt.get('checked', False):
                weight = lt.get('value', 0)
                checked_types.append((key, weight, leave_type_labels.get(key, key)))

        checked_types.sort(key=lambda x: x[1], reverse=True)

        if checked_types:
            type_phrases = []
            for i, (key, weight, label) in enumerate(checked_types):
                if i == 0:
                    type_phrases.append(f"mainly {label}")
                elif i == 1:
                    type_phrases.append(f"but also some {label}")
                elif i == 2:
                    type_phrases.append(f"and to a lesser extent {label}")
                else:
                    type_phrases.append(f"and occasionally {label}")

            if len(type_phrases) == 1:
                type_str = type_phrases[0]
            else:
                type_str = ", ".join(type_phrases)

            parts.append(f"Over the past two years, {p['subj_l']} {p['have']} taken approximately {leaves} {leave_type_label} leave{'s' if leaves != '1' else ''} {frequency}, averaging {duration} per leave, engaging in {type_str}.")

        # Other leave types
        other_types = []
        if state_dict.get('medical', False):
            other_types.append("medical appointments")
        if state_dict.get('court', False):
            other_types.append("court appearances")
        if state_dict.get('compassionate', False):
            other_types.append("compassionate visits")

        if other_types:
            if len(other_types) == 1:
                parts.append(f"{p['subj']} {p['have']} also taken leave for {other_types[0]}.")
            else:
                parts.append(f"{p['subj']} {p['have']} also taken leave for {', '.join(other_types[:-1])} and {other_types[-1]}.")

        # Leave suspension
        if state_dict.get('suspended_yes', False):
            details = state_dict.get('suspension_details', '').strip()
            if details:
                parts.append(f"{p['pos'].capitalize()} leave has been suspended in the past. {details}")
            else:
                parts.append(f"{p['pos'].capitalize()} leave has been suspended in the past.")
        elif state_dict.get('suspended_no', False):
            parts.append(f"{p['pos'].capitalize()} leave has never been suspended.")

        return " ".join(parts)

    def _update_leave_report_preview(self):
        """Update the leave report preview with generated narrative."""
        if not hasattr(self, 'leave_report_preview'):
            return

        # Initialize state dicts if they don't exist
        if not hasattr(self, '_leave_escorted_state'):
            self._leave_escorted_state = {}
        if not hasattr(self, '_leave_unescorted_state'):
            self._leave_unescorted_state = {}

        # Save current state to appropriate dict
        is_escorted = hasattr(self, 'leave_escorted_radio') and self.leave_escorted_radio.isChecked()
        if hasattr(self, 'leaves_dropdown'):  # Only save if popup has been built
            if is_escorted:
                self._save_leave_state(self._leave_escorted_state)
            else:
                self._save_leave_state(self._leave_unescorted_state)

        # Generate both outputs
        escorted_text = self._generate_leave_report_from_state(self._leave_escorted_state, "escorted")
        unescorted_text = self._generate_leave_report_from_state(self._leave_unescorted_state, "unescorted")

        result_parts = []
        if escorted_text:
            result_parts.append(f"ESCORTED LEAVE:\n{escorted_text}")
        if unescorted_text:
            result_parts.append(f"UNESCORTED LEAVE:\n{unescorted_text}")

        result = "\n\n".join(result_parts) if result_parts else "(No content yet...)"
        self.leave_report_preview.setPlainText(result)

    def _search_leave_evidence(self):
        """Search uploaded notes for evidence of leave taken and suspensions."""
        # Leave search terms
        leave_terms = [
            "took leave", "community leave", "went on leave", "returned from leave",
            "S17 leave", "engaged in S17", "ground leave", "escorted leave",
            "unescorted leave", "overnight leave", "local leave", "extended leave",
            "leave to", "leave was", "leave today", "on leave", "from leave"
        ]

        # Suspension search terms
        suspension_terms = [
            "leave suspended", "suspension of leave", "leave was suspended",
            "suspended leave", "leave revoked", "leave cancelled", "leave stopped",
            "leave withdrawn", "no leave", "leave not permitted"
        ]

        # Check if we have uploaded notes (from main app)
        main_window = self.window()
        notes_text = ""

        # Try to get notes from various sources
        if hasattr(main_window, 'uploaded_notes'):
            notes_text = main_window.uploaded_notes
        elif hasattr(self, 'leave_taken_text'):
            # Allow manual paste for now
            pass

        if not notes_text:
            self.leave_taken_text.setPlainText("No notes available. Please upload clinical notes via Data Extractor first.")
            self.leave_suspension_text.setPlainText("No notes available.")
            return

        # Parse notes and search (simplified - would need proper date parsing in production)
        leave_findings = []
        suspension_findings = []

        lines = notes_text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()

            # Check for leave evidence
            for term in leave_terms:
                if term.lower() in line_lower:
                    # Get context (2 lines before/after)
                    start = max(0, i - 1)
                    end = min(len(lines), i + 2)
                    context = '\n'.join(lines[start:end])
                    leave_findings.append(f"[{term}]: {context[:200]}...")
                    break

            # Check for suspension evidence
            for term in suspension_terms:
                if term.lower() in line_lower:
                    start = max(0, i - 1)
                    end = min(len(lines), i + 2)
                    context = '\n'.join(lines[start:end])
                    suspension_findings.append(f"[{term}]: {context[:200]}...")
                    break

        # Display findings
        if leave_findings:
            self.leave_taken_text.setPlainText('\n\n'.join(leave_findings[:10]))  # Limit to 10
        else:
            self.leave_taken_text.setPlainText("No evidence of leave found in notes.")

        if suspension_findings:
            self.leave_suspension_text.setPlainText('\n\n'.join(suspension_findings[:5]))
        else:
            self.leave_suspension_text.setPlainText("No evidence of leave suspension found in notes.")

    def _upload_leave_file(self):
        """Upload a file for leave report data extraction."""
        from data_extractor_popup import DataExtractorPopup

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Upload Leave Data File",
            "",
            "All Supported Files (*.pdf *.docx *.doc *.txt *.rtf *.xls *.xlsx);;PDF Files (*.pdf);;Word Documents (*.docx *.doc);;Excel Files (*.xls *.xlsx);;All Files (*)"
        )
        if not file_path:
            return

        # Create or get data extractor
        if not hasattr(self, '_leave_data_extractor') or self._leave_data_extractor is None:
            self._leave_data_extractor = DataExtractorPopup(parent=self)
            self._leave_data_extractor.data_extracted.connect(self._on_leave_data_extracted)

        # Import the file
        self._leave_data_extractor.load_file(file_path)
        self._leave_data_extractor.show()
        self._leave_data_extractor.raise_()
        self._leave_data_extractor.activateWindow()

    def _view_leave_data(self):
        """View the data extractor with previously loaded leave data."""
        from data_extractor_popup import DataExtractorPopup

        # Create or get data extractor
        if not hasattr(self, '_leave_data_extractor') or self._leave_data_extractor is None:
            self._leave_data_extractor = DataExtractorPopup(parent=self)
            self._leave_data_extractor.data_extracted.connect(self._on_leave_data_extracted)

        self._leave_data_extractor.show()
        self._leave_data_extractor.raise_()
        self._leave_data_extractor.activateWindow()

    def _on_leave_data_extracted(self, panel_data: dict):
        """Handle extracted data specifically for leave report section."""
        # Search for leave-related entries in the extracted data
        leave_text_parts = []
        suspension_text_parts = []

        # Get raw notes if available
        raw_notes = []
        if hasattr(self, '_leave_data_extractor') and self._leave_data_extractor:
            raw_notes = getattr(self._leave_data_extractor, 'notes', [])
            if raw_notes:
                notes_text = "\n".join([str(n.get('body', '')) for n in raw_notes if isinstance(n, dict)])
                self._search_leave_in_text(notes_text, leave_text_parts, suspension_text_parts)

        # Also check LEAVE category from extracted data
        if "LEAVE" in panel_data:
            notes = panel_data["LEAVE"]
            if notes:
                text = "\n\n".join(notes) if isinstance(notes, list) else str(notes)
                leave_text_parts.append(text)

        # Update the display
        if leave_text_parts:
            self.leave_taken_text.setPlainText('\n\n'.join(leave_text_parts))
        if suspension_text_parts:
            self.leave_suspension_text.setPlainText('\n\n'.join(suspension_text_parts))

        # Populate the Imported Notes section with filtered raw notes
        if raw_notes:
            self._populate_leave_imports(raw_notes)
            if hasattr(self, 'leave_import_section'):
                self.leave_import_section.setVisible(True)

    def _search_leave_in_text(self, text: str, leave_parts: list, suspension_parts: list):
        """Search for leave and suspension evidence in text."""
        leave_terms = ["took leave", "community leave", "went on leave", "returned from leave",
                      "S17 leave", "engaged in S17", "ground leave", "escorted leave",
                      "unescorted leave", "overnight leave", "local leave", "extended leave"]
        suspension_terms = ["leave suspended", "suspension of leave", "leave was suspended",
                           "suspended leave", "leave revoked", "leave cancelled"]

        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()
            for term in leave_terms:
                if term.lower() in line_lower:
                    start = max(0, i - 1)
                    end = min(len(lines), i + 2)
                    context = '\n'.join(lines[start:end])
                    leave_parts.append(f"[{term}]: {context[:200]}...")
                    break
            for term in suspension_terms:
                if term.lower() in line_lower:
                    start = max(0, i - 1)
                    end = min(len(lines), i + 2)
                    context = '\n'.join(lines[start:end])
                    suspension_parts.append(f"[{term}]: {context[:200]}...")
                    break

    def _populate_leave_imports(self, raw_notes: list):
        """
        Populate the Imported Notes section with filtered notes from last 2 years.
        Filters for leave-related categories and shows relevant entries with checkboxes.
        Like ASR section 6/7 with clickable labels and +/- expand/collapse.
        """
        import re
        from datetime import datetime, timedelta

        if not raw_notes:
            print("[MOJ-LEAVE] No raw notes available for leave imports")
            self._clear_leave_imports()
            return

        # Leave categories with keywords and colors (specific terms only)
        LEAVE_CATEGORIES = {
            "Ground Leave": [
                "ground leave", "grounds leave", "hospital grounds", "within grounds"
            ],
            "Local Community": [
                "local community leave", "local community", "local leave"
            ],
            "Community Leave": [
                "community leave", "s17 leave", "section 17 leave", "took leave today",
                "went on leave", "returned from leave", "engaged in s17 leave",
                "leave taken", "leave was successful", "leave uneventful"
            ],
            "Extended Leave": [
                "extended leave", "extended community leave", "overnight leave"
            ],
            "Escorted": [
                "escorted leave", "with escort", "staff escort", "nurse escort",
                "1:1 escort", "2:1 escort", "escorted by"
            ],
            "Unescorted": [
                "unescorted leave", "independent leave", "without escort", "unescorted community"
            ],
            "Suspension": [
                "leave suspended", "suspension of leave", "leave cancelled", "leave revoked",
                "leave stopped", "leave withdrawn", "leave incident"
            ]
        }

        CATEGORY_COLORS = {
            "Ground Leave": "#059669",      # Green
            "Local Community": "#0891b2",   # Cyan
            "Community Leave": "#3b82f6",   # Blue
            "Extended Leave": "#7c3aed",    # Purple
            "Escorted": "#d97706",          # Amber
            "Unescorted": "#dc2626",        # Red
            "Suspension": "#be185d"         # Pink
        }

        # Store for filtering
        self._leave_categories = LEAVE_CATEGORIES
        self._leave_category_colors = CATEGORY_COLORS

        # Parse date helper
        def parse_note_date(date_val):
            if isinstance(date_val, datetime):
                return date_val
            if not date_val:
                return None
            date_str = str(date_val).strip()
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"]:
                try:
                    return datetime.strptime(date_str.split()[0] if ' ' in date_str else date_str, fmt.split()[0])
                except:
                    pass
            return None

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            if 11 <= day <= 13:
                suffix = "th"
            else:
                suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def get_matching_categories(text: str) -> list:
            """Return list of categories that match the text."""
            if not text:
                return []
            text_lower = text.lower()
            matched = []
            for cat, keywords in LEAVE_CATEGORIES.items():
                if any(kw in text_lower for kw in keywords):
                    matched.append(cat)
            return matched

        def extract_snippet(text: str, categories: list) -> str:
            """Extract first 2 lines + lines containing filter keywords."""
            if not text:
                return ""
            lines = text.split('\n')
            lines = [l.strip() for l in lines if l.strip()]
            if not lines:
                return text[:300]

            result_lines = lines[:2]

            # Get all keywords for matched categories
            all_keywords = []
            for cat in categories:
                all_keywords.extend(LEAVE_CATEGORIES.get(cat, []))

            # Find lines containing keywords (after first 2)
            for i, line in enumerate(lines[2:], start=2):
                line_lower = line.lower()
                if any(kw in line_lower for kw in all_keywords):
                    if result_lines and i > len(result_lines):
                        if result_lines[-1] != "...":
                            result_lines.append("...")
                    result_lines.append(line)

            return '\n'.join(result_lines)

        # Find the most recent note date
        all_dates = []
        for n in raw_notes:
            dt = parse_note_date(n.get("date") or n.get("datetime"))
            if dt:
                all_dates.append(dt)

        if not all_dates:
            print("[MOJ-LEAVE] No parseable dates in notes")
            self._clear_leave_imports()
            return

        most_recent = max(all_dates)
        two_years_cutoff = most_recent - timedelta(days=730)

        print(f"[MOJ-LEAVE] Most recent note: {most_recent.strftime('%d/%m/%Y')}")
        print(f"[MOJ-LEAVE] 2-year cutoff: {two_years_cutoff.strftime('%d/%m/%Y')}")

        # Filter and categorize notes
        categorized = []
        seen_texts = set()

        for n in raw_notes:
            note_date = parse_note_date(n.get("date") or n.get("datetime"))
            if not note_date:
                continue

            if note_date < two_years_cutoff:
                continue

            full_text = n.get("body", "") or n.get("text", "") or n.get("content", "")
            if not full_text:
                continue

            categories = get_matching_categories(full_text)
            if not categories:
                continue

            text_sig = full_text[:200].strip()
            if text_sig in seen_texts:
                continue
            seen_texts.add(text_sig)

            snippet = extract_snippet(full_text, categories)

            categorized.append({
                "text": full_text,
                "snippet": snippet,
                "date": n.get("date") or n.get("datetime"),
                "date_obj": note_date,
                "categories": categories
            })

        # Sort by date descending
        categorized.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)

        print(f"[MOJ-LEAVE] Found {len(categorized)} categorized leave entries in last 2 years")

        # Limit to prevent performance issues (show most recent 150)
        MAX_DISPLAY = 150
        if len(categorized) > MAX_DISPLAY:
            print(f"[MOJ-LEAVE] Limiting display to {MAX_DISPLAY} most recent entries")
            categorized = categorized[:MAX_DISPLAY]

        # Store for filtering
        self._leave_all_categorized = categorized

        # Display entries
        self._display_leave_entries(categorized, format_date_nice, None)

    def _display_leave_entries(self, categorized: list, format_date_nice, filter_category: str = None):
        """Display leave entries with optional category filter."""
        import re
        import html

        if not hasattr(self, 'leave_import_layout') or not self.leave_import_layout:
            return

        CATEGORY_COLORS = self._leave_category_colors
        LEAVE_CATEGORIES = self._leave_categories

        def highlight_keywords(text, matched_categories):
            """Highlight keywords in text that triggered the category matches."""
            escaped = html.escape(text)
            keywords_to_highlight = []
            for cat in matched_categories:
                keywords_to_highlight.extend(LEAVE_CATEGORIES.get(cat, []))
            keywords_to_highlight = sorted(set(keywords_to_highlight), key=len, reverse=True)
            for kw in keywords_to_highlight:
                color = CATEGORY_COLORS.get(matched_categories[0], "#fef08a") if matched_categories else "#fef08a"
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m: f'<span style="background-color: {color}; color: white; padding: 1px 2px; border-radius: 2px; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            return escaped.replace('\n', '<br>')

        # Clear existing content
        while self.leave_import_layout.count():
            item = self.leave_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                while item.layout().count():
                    sub_item = item.layout().takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()

        self.leave_import_checkboxes = []

        # Filter if category specified
        if filter_category:
            display_entries = [(e, e["categories"]) for e in categorized if filter_category in e.get("categories", [])]
        else:
            display_entries = [(e, e["categories"]) for e in categorized]

        # Update subtitle
        if hasattr(self, 'leave_import_subtitle'):
            self.leave_import_subtitle.setText(f"Showing {len(display_entries)} leave entries from last 2 years - tick to include in preview")

        if not display_entries:
            placeholder = QLabel("No leave entries found. Use Import Data to load clinical notes.")
            placeholder.setStyleSheet("color: #92400e; font-style: italic; font-size: 11px; background: transparent;")
            self.leave_import_layout.addWidget(placeholder)
            return

        # Filter header row (only show when filter is active) - red box style
        if filter_category:
            filter_frame = QFrame()
            filter_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #dc2626; border-radius: 4px; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(8, 4, 8, 4)
            filter_row.setSpacing(8)

            filter_label = QLabel(f"Filtered by:")
            filter_label.setStyleSheet("font-size: 10px; color: #991b1b; background: transparent;")
            filter_row.addWidget(filter_label)

            color = CATEGORY_COLORS.get(filter_category, "#6b7280")
            cat_label = QLabel(filter_category)
            cat_label.setStyleSheet(f"font-size: 10px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(cat_label)

            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 9px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 3px; padding: 2px 6px;
                }
                QPushButton:hover { background: #fee2e2; }
            """)
            remove_btn.clicked.connect(lambda: self._filter_leave_by_category(None))
            filter_row.addWidget(remove_btn)

            self.leave_import_layout.addWidget(filter_frame)

        # Select All / Deselect All buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(6)

        select_all_btn = QPushButton("Select All")
        select_all_btn.setStyleSheet("""
            QPushButton { background: #d97706; color: white; border: none; border-radius: 4px; padding: 4px 10px; font-size: 10px; font-weight: 600; }
            QPushButton:hover { background: #b45309; }
        """)
        select_all_btn.clicked.connect(self._select_all_leave_imports)
        btn_row.addWidget(select_all_btn)

        deselect_all_btn = QPushButton("Deselect All")
        deselect_all_btn.setStyleSheet("""
            QPushButton { background: #92400e; color: white; border: none; border-radius: 4px; padding: 4px 10px; font-size: 10px; font-weight: 600; }
            QPushButton:hover { background: #78350f; }
        """)
        deselect_all_btn.clicked.connect(self._deselect_all_leave_imports)
        btn_row.addWidget(deselect_all_btn)

        btn_row.addStretch()
        self.leave_import_layout.addLayout(btn_row)

        # Add each entry
        for entry, categories in display_entries:
            full_text = entry["text"]
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            display_text = entry.get("snippet", full_text[:150] + "..." if len(full_text) > 150 else full_text)

            entry_frame = QFrame()
            border_color = CATEGORY_COLORS.get(categories[0], "#fcd34d") if categories else "#fcd34d"
            entry_frame.setStyleSheet(f"QFrame {{ background: white; border: 1px solid {border_color}; border-radius: 4px; }}")

            entry_layout_h = QHBoxLayout(entry_frame)
            entry_layout_h.setContentsMargins(6, 4, 6, 4)
            entry_layout_h.setSpacing(6)

            cb = QCheckBox()
            cb.setChecked(False)  # Not checked by default
            cb.setStyleSheet("QCheckBox { background: transparent; }")
            # Auto-update preview when checkbox state changes
            cb.stateChanged.connect(lambda state: self._update_leave_preview_from_imports())
            entry_layout_h.addWidget(cb)

            content_layout = QVBoxLayout()
            content_layout.setSpacing(2)

            # Header with tags and expand button
            header_layout = QHBoxLayout()
            header_layout.setSpacing(4)

            if categories:
                for cat in categories:
                    tag = QPushButton(cat)
                    tag.setCursor(Qt.CursorShape.PointingHandCursor)
                    color = CATEGORY_COLORS.get(cat, "#6b7280")
                    tag.setStyleSheet(f"""
                        QPushButton {{
                            font-size: 9px; font-weight: 600; color: white;
                            background: {color}; padding: 1px 4px; border-radius: 3px; border: none;
                        }}
                        QPushButton:hover {{ background: {color}; opacity: 0.8; }}
                    """)
                    tag.clicked.connect(lambda checked, c=cat: self._filter_leave_by_category(c))
                    header_layout.addWidget(tag)

            header_layout.addStretch()

            # Expand/collapse button - starts collapsed (shows +)
            expand_btn = QPushButton("+")
            expand_btn.setFixedSize(20, 20)
            expand_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            expand_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; font-weight: 700; color: #6b7280;
                    background: #f3f4f6; border: 1px solid #d1d5db; border-radius: 4px;
                }
                QPushButton:hover { background: #e5e7eb; color: #374151; }
            """)
            header_layout.addWidget(expand_btn)
            content_layout.addLayout(header_layout)

            # Date
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            if date_str:
                date_lbl = QLabel(date_str)
                date_lbl.setStyleSheet("font-size: 10px; font-weight: 600; color: #92400e; background: transparent;")
                content_layout.addWidget(date_lbl)

            # Text label (snippet by default)
            text_lbl = QLabel(display_text[:200] + "..." if len(display_text) > 200 else display_text)
            text_lbl.setWordWrap(True)
            text_lbl.setStyleSheet("font-size: 10px; color: #374151; background: transparent;")
            text_lbl.setTextFormat(Qt.TextFormat.PlainText)
            content_layout.addWidget(text_lbl)

            # Store state for expand/collapse - starts collapsed
            entry_frame._is_expanded = False
            entry_frame._snippet_text = display_text[:200] + "..." if len(display_text) > 200 else display_text
            entry_frame._full_highlighted = highlight_keywords(full_text, categories)
            entry_frame._text_lbl = text_lbl
            entry_frame._expand_btn = expand_btn

            def make_toggle(ef, tl, eb):
                def toggle_expand():
                    if ef._is_expanded:
                        # Collapse: show snippet, button shows +
                        tl.setTextFormat(Qt.TextFormat.PlainText)
                        tl.setText(ef._snippet_text)
                        eb.setText("+")
                        ef._is_expanded = False
                    else:
                        # Expand: show full highlighted, button shows −
                        tl.setTextFormat(Qt.TextFormat.RichText)
                        tl.setText(ef._full_highlighted)
                        eb.setText("−")
                        ef._is_expanded = True
                return toggle_expand

            expand_btn.clicked.connect(make_toggle(entry_frame, text_lbl, expand_btn))

            entry_layout_h.addLayout(content_layout, 1)

            self.leave_import_layout.addWidget(entry_frame)
            self.leave_import_checkboxes.append((cb, entry))

    def _update_leave_preview_from_imports(self):
        """Auto-update the leave report preview when import checkboxes change."""
        if not hasattr(self, 'leave_import_checkboxes') or not self.leave_import_checkboxes:
            return

        # Collect checked entries
        texts = []
        for checkbox, entry in self.leave_import_checkboxes:
            if checkbox.isChecked():
                text = entry.get("text", "").strip()
                date_obj = entry.get("date_obj")
                if text:
                    if date_obj:
                        date_str = date_obj.strftime("%d/%m/%Y")
                        texts.append(f"[{date_str}] {text}")
                    else:
                        texts.append(text)

        # Update preview - append to generated content
        if hasattr(self, 'leave_report_preview'):
            # First regenerate the base preview from input fields
            self._update_leave_report_preview()

            # Then append imported notes if any are selected
            if texts:
                current = self.leave_report_preview.toPlainText()
                imported_section = "\n\n--- Imported Leave Notes ---\n" + "\n\n".join(texts)
                if current:
                    self.leave_report_preview.setPlainText(current + imported_section)
                else:
                    self.leave_report_preview.setPlainText(imported_section.strip())

    def _filter_leave_by_category(self, category: str):
        """Filter leave entries by category."""
        self._leave_current_filter = category

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        self._display_leave_entries(self._leave_all_categorized, format_date_nice, category)

    def _select_all_leave_imports(self):
        """Select all leave import checkboxes."""
        for checkbox, _ in self.leave_import_checkboxes:
            checkbox.setChecked(True)

    def _deselect_all_leave_imports(self):
        """Deselect all leave import checkboxes."""
        for checkbox, _ in self.leave_import_checkboxes:
            checkbox.setChecked(False)

    def _populate_leave_imports_from_items(self, items: list):
        """Populate leave imports from pre-categorized items (from data extractor)."""
        from datetime import datetime

        if not items:
            self._clear_leave_imports()
            return

        if not hasattr(self, 'leave_import_layout') or not self.leave_import_layout:
            return

        # Clear existing
        self._clear_leave_imports()

        # Limit items to prevent performance issues
        MAX_DISPLAY = 100
        display_items = items[:MAX_DISPLAY] if len(items) > MAX_DISPLAY else items
        print(f"[MOJ-LEAVE] Displaying {len(display_items)} leave items (of {len(items)} total)")

        # Update subtitle
        if hasattr(self, 'leave_import_subtitle'):
            self.leave_import_subtitle.setText(f"Imported entries ({len(display_items)} shown)")

        self.leave_import_checkboxes = []

        for item in display_items:
            text = item.get("text", "").strip()
            if not text:
                continue

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #fcd34d; border-radius: 4px; }")
            entry_layout = QHBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 4, 6, 4)
            entry_layout.setSpacing(6)

            # Checkbox
            cb = QCheckBox()
            cb.setStyleSheet("QCheckBox { background: transparent; }")
            cb.stateChanged.connect(self._update_leave_preview_from_imports)
            entry_layout.addWidget(cb)

            # Text preview (first 150 chars)
            preview = text[:150] + "..." if len(text) > 150 else text
            text_label = QLabel(preview)
            text_label.setWordWrap(True)
            text_label.setStyleSheet("font-size: 11px; color: #78350f; background: transparent;")
            entry_layout.addWidget(text_label, 1)

            self.leave_import_layout.addWidget(entry_frame)
            self.leave_import_checkboxes.append((cb, item))

    def _populate_psych_imports_from_items(self, items: list):
        """Populate psych imports from pre-categorized items (from data extractor)."""
        from datetime import datetime

        if not items:
            self._clear_psych_imports()
            return

        if not hasattr(self, 'psych_import_layout') or not self.psych_import_layout:
            return

        # Clear existing
        self._clear_psych_imports()

        # Limit items to prevent performance issues
        MAX_DISPLAY = 100
        display_items = items[:MAX_DISPLAY] if len(items) > MAX_DISPLAY else items
        print(f"[MOJ-LEAVE] Displaying {len(display_items)} psych items (of {len(items)} total)")

        # Update subtitle
        if hasattr(self, 'psych_import_subtitle'):
            self.psych_import_subtitle.setText(f"Imported entries ({len(display_items)} shown)")

        self.psych_import_checkboxes = []

        for item in display_items:
            text = item.get("text", "").strip()
            if not text:
                continue

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #fcd34d; border-radius: 4px; }")
            entry_layout = QHBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 4, 6, 4)
            entry_layout.setSpacing(6)

            # Checkbox
            cb = QCheckBox()
            cb.setStyleSheet("QCheckBox { background: transparent; }")
            cb.stateChanged.connect(self._update_psych_preview_from_imports)
            entry_layout.addWidget(cb)

            # Text preview (first 150 chars)
            preview = text[:150] + "..." if len(text) > 150 else text
            text_label = QLabel(preview)
            text_label.setWordWrap(True)
            text_label.setStyleSheet("font-size: 11px; color: #78350f; background: transparent;")
            entry_layout.addWidget(text_label, 1)

            self.psych_import_layout.addWidget(entry_frame)
            self.psych_import_checkboxes.append((cb, item))

    def _clear_leave_imports(self):
        """Clear the imported leave entries."""
        self.leave_imported_entries = []
        self.leave_import_checkboxes = []
        self._leave_all_categorized = []
        self._leave_current_filter = None

        if not hasattr(self, 'leave_import_layout') or not self.leave_import_layout:
            return

        while self.leave_import_layout.count():
            item = self.leave_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                while item.layout().count():
                    sub_item = item.layout().takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()

        placeholder = QLabel("No entries found. Use Import Data to load clinical notes.")
        placeholder.setStyleSheet("color: #92400e; font-style: italic; font-size: 11px; background: transparent;")
        self.leave_import_layout.addWidget(placeholder)

        if hasattr(self, 'leave_import_subtitle') and self.leave_import_subtitle:
            self.leave_import_subtitle.setText("Leave evidence from notes (last 2 years)")

    def _send_leave_imports_to_preview(self):
        """Send selected (checked) leave entries to the preview text area."""
        if not hasattr(self, 'leave_import_checkboxes') or not self.leave_import_checkboxes:
            QMessageBox.information(self, "No Entries", "No leave entries available to send.")
            return

        texts = []
        for checkbox, entry in self.leave_import_checkboxes:
            if checkbox.isChecked():
                text = entry.get("text", "").strip()
                date_obj = entry.get("date_obj")
                if text:
                    if date_obj:
                        date_str = date_obj.strftime("%d/%m/%Y")
                        texts.append(f"[{date_str}] {text}")
                    else:
                        texts.append(text)

        if not texts:
            QMessageBox.information(self, "No Selection", "Please tick at least one entry to send to the preview.")
            return

        combined = "\n\n".join(texts)

        if hasattr(self, 'leave_report_preview'):
            current = self.leave_report_preview.toPlainText()
            if current:
                self.leave_report_preview.setPlainText(current + "\n\n--- Imported Leave Notes ---\n" + combined)
            else:
                self.leave_report_preview.setPlainText(combined)

        QMessageBox.information(self, "Sent", f"{len(texts)} leave entries sent to preview.")

    def _create_procedures_popup(self) -> QWidget:
        """Create procedures popup with escorts, transport, handcuffs, exclusion zones, etc."""
        popup = QWidget()
        popup.setStyleSheet("background: #f9fafb;")
        main_layout = QVBoxLayout(popup)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # === PREVIEW SECTION (full black background, white text) ===
        preview_container = QFrame()
        preview_container.setStyleSheet("QFrame { background: black; border-radius: 8px; }")
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 10, 12, 10)
        preview_layout.setSpacing(8)

        # Header row with title and Send to Card button
        preview_header = QHBoxLayout()
        preview_title = QLabel("Preview")
        preview_title.setStyleSheet("font-size: 12px; font-weight: 600; color: white; background: transparent;")
        preview_header.addWidget(preview_title)
        preview_header.addStretch()

        send_btn = QPushButton("Send to Card")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #dc2626;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 4px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #b91c1c; }
        """)
        send_btn.clicked.connect(lambda: self._send_to_card("procedures"))
        preview_header.addWidget(send_btn)
        preview_layout.addLayout(preview_header)

        self.leave_procedures = QTextEdit()
        self.leave_procedures.setPlaceholderText("Management procedures will be generated based on selections below...")
        self.leave_procedures.setMinimumHeight(100)
        self.leave_procedures.setMaximumHeight(140)
        self.leave_procedures.setStyleSheet("""
            QTextEdit {
                background: black;
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px;
                font-size: 11px;
            }
            QTextEdit::placeholder { color: #6b7280; }
        """)
        preview_layout.addWidget(self.leave_procedures)
        main_layout.addWidget(preview_container)

        main_layout.addSpacing(12)

        # === INPUT SECTIONS (scrollable) ===
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll_area.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 8, 0)
        scroll_layout.setSpacing(8)

        # Label style without borders
        label_style = "font-size: 11px; color: #374151; background: transparent; border: none;"
        cb_style = "font-size: 11px; color: #374151; background: transparent; border: none;"

        # === SECTION 1: Escorts/Transport ===
        self.section1_collapsed = False
        section1_header = QHBoxLayout()
        self.section1_btn = QPushButton("−")
        self.section1_btn.setFixedSize(20, 20)
        self.section1_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.section1_btn.setStyleSheet("""
            QPushButton { font-size: 14px; font-weight: 700; color: #374151; background: #e5e7eb; border: none; border-radius: 4px; }
            QPushButton:hover { background: #d1d5db; }
        """)
        section1_header.addWidget(self.section1_btn)
        section1_title = QLabel("Escorts / Transport")
        section1_title.setStyleSheet("font-size: 12px; font-weight: 600; color: #1f2937; background: transparent;")
        section1_header.addWidget(section1_title)
        section1_header.addStretch()
        scroll_layout.addLayout(section1_header)

        self.section1_content = QWidget()
        self.section1_content.setStyleSheet("background: transparent;")
        section1_layout = QVBoxLayout(self.section1_content)
        section1_layout.setContentsMargins(28, 4, 0, 8)
        section1_layout.setSpacing(6)

        # Escorts
        escorts_row = QHBoxLayout()
        escorts_row.setSpacing(8)
        escorts_lbl = QLabel("Escorts:")
        escorts_lbl.setStyleSheet(label_style)
        escorts_row.addWidget(escorts_lbl)
        self.escorts_dropdown = QComboBox()
        self.escorts_dropdown.addItems(["Select...", "1", "2", "3"])
        self.escorts_dropdown.setFixedWidth(70)
        self.escorts_dropdown.setStyleSheet("font-size: 11px;")
        self.escorts_dropdown.currentTextChanged.connect(self._update_section_3h_text)
        escorts_row.addWidget(self.escorts_dropdown)
        escorts_row.addStretch()
        section1_layout.addLayout(escorts_row)

        # Transport
        transport_lbl = QLabel("Transport:")
        transport_lbl.setStyleSheet(label_style)
        section1_layout.addWidget(transport_lbl)
        transport_row = QHBoxLayout()
        transport_row.setSpacing(12)
        self.transport_checkboxes = {}
        for key, label in [("secure", "Secure"), ("hospital", "Hospital"), ("taxi", "Taxi"), ("public", "Public")]:
            cb = QCheckBox(label)
            cb.setStyleSheet(cb_style)
            cb.stateChanged.connect(self._update_section_3h_text)
            self.transport_checkboxes[key] = cb
            transport_row.addWidget(cb)
        transport_row.addStretch()
        section1_layout.addLayout(transport_row)

        # Handcuffs
        self.handcuffs_cb = QCheckBox("Handcuffs to be carried")
        self.handcuffs_cb.setStyleSheet(cb_style)
        self.handcuffs_cb.stateChanged.connect(self._update_section_3h_text)
        section1_layout.addWidget(self.handcuffs_cb)

        # Exclusion Zone
        excl_row = QHBoxLayout()
        excl_row.setSpacing(8)
        excl_lbl = QLabel("Exclusion Zone:")
        excl_lbl.setStyleSheet(label_style)
        excl_row.addWidget(excl_lbl)
        self.exclusion_group = QButtonGroup(self)
        self.exclusion_yes = QRadioButton("Yes")
        self.exclusion_yes.setStyleSheet(cb_style)
        self.exclusion_na = QRadioButton("N/A")
        self.exclusion_na.setStyleSheet(cb_style)
        self.exclusion_na.setChecked(True)
        self.exclusion_group.addButton(self.exclusion_yes)
        self.exclusion_group.addButton(self.exclusion_na)
        excl_row.addWidget(self.exclusion_yes)
        excl_row.addWidget(self.exclusion_na)
        excl_row.addStretch()
        section1_layout.addLayout(excl_row)

        self.exclusion_details = QLineEdit()
        self.exclusion_details.setPlaceholderText("Exclusion zone details...")
        self.exclusion_details.setStyleSheet("font-size: 11px; padding: 4px 8px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.exclusion_details.setVisible(False)
        self.exclusion_yes.toggled.connect(lambda checked: self.exclusion_details.setVisible(checked))
        self.exclusion_yes.toggled.connect(self._update_section_3h_text)
        self.exclusion_na.toggled.connect(self._update_section_3h_text)
        self.exclusion_details.textChanged.connect(self._update_section_3h_text)
        section1_layout.addWidget(self.exclusion_details)

        scroll_layout.addWidget(self.section1_content)

        def toggle_section1():
            self.section1_collapsed = not self.section1_collapsed
            self.section1_content.setVisible(not self.section1_collapsed)
            self.section1_btn.setText("+" if self.section1_collapsed else "−")
        self.section1_btn.clicked.connect(toggle_section1)

        # === SECTION 2: Public Protection Pre-leave ===
        self.section2_collapsed = False
        section2_header = QHBoxLayout()
        self.section2_btn = QPushButton("−")
        self.section2_btn.setFixedSize(20, 20)
        self.section2_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.section2_btn.setStyleSheet("""
            QPushButton { font-size: 14px; font-weight: 700; color: #374151; background: #e5e7eb; border: none; border-radius: 4px; }
            QPushButton:hover { background: #d1d5db; }
        """)
        section2_header.addWidget(self.section2_btn)
        section2_title = QLabel("Public Protection - Pre-leave")
        section2_title.setStyleSheet("font-size: 12px; font-weight: 600; color: #1f2937; background: transparent;")
        section2_header.addWidget(section2_title)
        section2_header.addStretch()
        scroll_layout.addLayout(section2_header)

        self.section2_content = QWidget()
        self.section2_content.setStyleSheet("background: transparent;")
        section2_layout = QVBoxLayout(self.section2_content)
        section2_layout.setContentsMargins(28, 4, 0, 8)
        section2_layout.setSpacing(4)

        self.pre_leave_checkboxes = {}
        pre_leave_options = [
            ("risk_free", "Risk free > 24 hours"),
            ("mental_state", "Mental state assessment prior to leave"),
            ("escorts_confirmed", "Escorts confirmed as known to patient"),
            ("no_drugs", "No permission for drug and alcohol use"),
            ("timings", "Timings monitored"),
        ]
        for key, label in pre_leave_options:
            cb = QCheckBox(label)
            cb.setStyleSheet(cb_style)
            cb.stateChanged.connect(self._update_section_3h_text)
            self.pre_leave_checkboxes[key] = cb
            section2_layout.addWidget(cb)

        scroll_layout.addWidget(self.section2_content)

        def toggle_section2():
            self.section2_collapsed = not self.section2_collapsed
            self.section2_content.setVisible(not self.section2_collapsed)
            self.section2_btn.setText("+" if self.section2_collapsed else "−")
        self.section2_btn.clicked.connect(toggle_section2)

        # === SECTION 3: Public Protection On Return ===
        self.section3_collapsed = False
        section3_header = QHBoxLayout()
        self.section3_btn = QPushButton("−")
        self.section3_btn.setFixedSize(20, 20)
        self.section3_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.section3_btn.setStyleSheet("""
            QPushButton { font-size: 14px; font-weight: 700; color: #374151; background: #e5e7eb; border: none; border-radius: 4px; }
            QPushButton:hover { background: #d1d5db; }
        """)
        section3_header.addWidget(self.section3_btn)
        section3_title = QLabel("Public Protection - On Return")
        section3_title.setStyleSheet("font-size: 12px; font-weight: 600; color: #1f2937; background: transparent;")
        section3_header.addWidget(section3_title)
        section3_header.addStretch()
        scroll_layout.addLayout(section3_header)

        self.section3_content = QWidget()
        self.section3_content.setStyleSheet("background: transparent;")
        section3_layout = QVBoxLayout(self.section3_content)
        section3_layout.setContentsMargins(28, 4, 0, 8)
        section3_layout.setSpacing(4)

        # On return checkboxes
        self.on_return_checkboxes = {}
        for key, label in [("search", "Search"), ("drug_testing", "Drug testing")]:
            cb = QCheckBox(label)
            cb.setStyleSheet(cb_style)
            cb.stateChanged.connect(self._update_section_3h_text)
            self.on_return_checkboxes[key] = cb
            section3_layout.addWidget(cb)

        # Breaches subsection
        breach_lbl = QLabel("Breaches:")
        breach_lbl.setStyleSheet(label_style + " margin-top: 6px;")
        section3_layout.addWidget(breach_lbl)

        self.breach_checkboxes = {}
        for key, label in [("suspension", "Suspension of leave"), ("inform_moj", "Inform MOJ")]:
            cb = QCheckBox(label)
            cb.setStyleSheet(cb_style)
            cb.stateChanged.connect(self._update_section_3h_text)
            self.breach_checkboxes[key] = cb
            section3_layout.addWidget(cb)

        # Specific to patient
        self.specific_to_patient = QCheckBox("I confirm measures are specific to this patient")
        self.specific_to_patient.setStyleSheet(cb_style + " margin-top: 6px;")
        self.specific_to_patient.stateChanged.connect(self._update_section_3h_text)
        section3_layout.addWidget(self.specific_to_patient)

        scroll_layout.addWidget(self.section3_content)

        def toggle_section3():
            self.section3_collapsed = not self.section3_collapsed
            self.section3_content.setVisible(not self.section3_collapsed)
            self.section3_btn.setText("+" if self.section3_collapsed else "−")
        self.section3_btn.clicked.connect(toggle_section3)

        scroll_layout.addStretch()
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area, 1)

        return popup

    def _update_section_3h_text(self):
        """Update the procedures text based on selections."""
        p = self._get_pronouns()
        parts = []

        # Escorts
        escorts = self.escorts_dropdown.currentText()
        if escorts and escorts != "Select...":
            parts.append(f"{p['subj']} will be accompanied by {escorts} escort{'s' if escorts != '1' else ''}")

        # Transport
        transport_selected = []
        for key, cb in self.transport_checkboxes.items():
            if cb.isChecked():
                if key == "public":
                    transport_selected.append("public transport")
                elif key == "hospital":
                    transport_selected.append("hospital transport")
                elif key == "secure":
                    transport_selected.append("secure transport")
                else:
                    transport_selected.append(key)
        if transport_selected:
            if len(transport_selected) == 1:
                parts.append(f"{p['subj']} will travel by {transport_selected[0]}")
            else:
                parts.append(f"{p['subj']} will travel by {' and by '.join(transport_selected)}")

        # Handcuffs
        if self.handcuffs_cb.isChecked():
            parts.append("Handcuffs will be carried")

        # Exclusion zone
        if self.exclusion_yes.isChecked():
            details = self.exclusion_details.text().strip()
            if details:
                parts.append(f"An exclusion zone applies: {details}")
            else:
                parts.append("An exclusion zone applies")
        elif self.exclusion_na.isChecked():
            parts.append("There is no exclusion zone")

        # Pre-leave checks
        pre_leave_parts = []
        if self.pre_leave_checkboxes["risk_free"].isChecked():
            pre_leave_parts.append(f"{p['subj_l']} must be risk free for more than 24 hours")
        if self.pre_leave_checkboxes["mental_state"].isChecked():
            pre_leave_parts.append(f"{p['pos_l']} mental state will be assessed prior to leave")
        if self.pre_leave_checkboxes["escorts_confirmed"].isChecked():
            pre_leave_parts.append("escorts will be confirmed as known to the patient")
        if self.pre_leave_checkboxes["timings"].isChecked():
            pre_leave_parts.append("all timings will be monitored")

        if pre_leave_parts:
            parts.append(f"Prior to leave, {' and '.join(pre_leave_parts)}")

        # No drugs/alcohol (separate sentence)
        if self.pre_leave_checkboxes["no_drugs"].isChecked():
            parts.append(f"{p['subj']} will not be permitted to take drugs or alcohol")

        # On return
        on_return_parts = []
        if self.on_return_checkboxes["search"].isChecked():
            on_return_parts.append(f"{p['subj_l']} will be searched")
        if self.on_return_checkboxes["drug_testing"].isChecked():
            on_return_parts.append(f"{p['subj_l']} will undergo drug testing")

        if on_return_parts:
            if len(on_return_parts) == 2:
                parts.append(f"On return {on_return_parts[0]} and {on_return_parts[1]}")
            else:
                parts.append(f"On return {on_return_parts[0]}")

        # Breaches
        breach_parts = []
        if self.breach_checkboxes["suspension"].isChecked():
            breach_parts.append("leave will be suspended")
        if self.breach_checkboxes["inform_moj"].isChecked():
            breach_parts.append("the MOJ will be informed")

        if breach_parts:
            parts.append(f"In the event of any breach {' and '.join(breach_parts)}")

        # Specific to patient
        if self.specific_to_patient.isChecked():
            parts.append("I can confirm the measures proposed are specifically defined for this patient")

        self.leave_procedures.setPlainText(". ".join(parts) + "." if parts else "")

    def _create_hospital_admissions_popup(self) -> QWidget:
        """Create hospital admissions popup with imported psychiatric history panel - like 3g style."""
        popup = QWidget()
        popup.setStyleSheet("background: #f9fafb;")
        main_layout = QVBoxLayout(popup)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # === PREVIEW SECTION (full black background, white text) ===
        preview_container = QFrame()
        preview_container.setStyleSheet("QFrame { background: black; border-radius: 8px; }")
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 10, 12, 10)
        preview_layout.setSpacing(8)

        # Header row with title and Send to Card button
        preview_header = QHBoxLayout()
        preview_title = QLabel("Preview - Past Psychiatric History")
        preview_title.setStyleSheet("font-size: 12px; font-weight: 600; color: white; background: transparent;")
        preview_header.addWidget(preview_title)
        preview_header.addStretch()

        send_btn = QPushButton("Send to Card")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #dc2626;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 4px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #b91c1c; }
        """)
        send_btn.clicked.connect(lambda: self._send_to_card("hospital_admissions"))
        preview_header.addWidget(send_btn)
        preview_layout.addLayout(preview_header)

        self.hospital_admissions = QTextEdit()
        self.hospital_admissions.setPlaceholderText("Past psychiatric history will be generated from selections below...")
        self.hospital_admissions.setMinimumHeight(100)
        self.hospital_admissions.setMaximumHeight(140)
        self.hospital_admissions.setStyleSheet("""
            QTextEdit {
                background: black;
                color: white;
                border: none;
                border-radius: 6px;
                padding: 8px;
                font-size: 11px;
            }
            QTextEdit::placeholder { color: #6b7280; }
        """)
        preview_layout.addWidget(self.hospital_admissions)
        main_layout.addWidget(preview_container)

        main_layout.addSpacing(12)

        # === IMPORTED NOTES SECTION (light orange like 3g) ===
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll_area.setStyleSheet("QScrollArea { border: none; background: transparent; }")

        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 8, 0)
        scroll_layout.setSpacing(8)

        # Imported Notes Section Header
        self.psych_import_section = QFrame()
        self.psych_import_section.setStyleSheet("QFrame { background: #fffbeb; border: 1px solid #fcd34d; border-radius: 8px; }")
        import_layout = QVBoxLayout(self.psych_import_section)
        import_layout.setContentsMargins(10, 10, 10, 10)
        import_layout.setSpacing(8)

        # Subtitle with count
        self.psych_import_subtitle = QLabel("Imported psychiatric history entries - tick to include in preview")
        self.psych_import_subtitle.setStyleSheet("font-size: 11px; font-weight: 600; color: #92400e; background: transparent;")
        import_layout.addWidget(self.psych_import_subtitle)

        # Scrollable content area for entries
        psych_import_scroll = QScrollArea()
        psych_import_scroll.setWidgetResizable(True)
        psych_import_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        psych_import_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        psych_import_scroll.setStyleSheet("QScrollArea { background: #fffbeb; border: none; }")
        psych_import_scroll.setMinimumHeight(200)

        # Content widget for entries
        self.psych_import_content = QWidget()
        self.psych_import_content.setStyleSheet("background: transparent;")
        self.psych_import_layout = QVBoxLayout(self.psych_import_content)
        self.psych_import_layout.setContentsMargins(4, 4, 4, 4)
        self.psych_import_layout.setSpacing(6)
        psych_import_scroll.setWidget(self.psych_import_content)

        # Placeholder
        self.psych_import_placeholder = QLabel("No entries found. Use Import Data to load clinical notes.")
        self.psych_import_placeholder.setStyleSheet("color: #92400e; font-style: italic; font-size: 11px; background: transparent;")
        self.psych_import_layout.addWidget(self.psych_import_placeholder)

        import_layout.addWidget(psych_import_scroll, 1)

        # Clear button
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(8)

        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet("""
            QPushButton { background: #fef3c7; color: #92400e; border: 1px solid #fcd34d; padding: 6px 12px; border-radius: 6px; font-weight: 500; font-size: 11px; }
            QPushButton:hover { background: #fde68a; }
        """)
        clear_btn.clicked.connect(self._clear_psych_imports)
        btn_layout.addWidget(clear_btn)

        btn_layout.addStretch()
        import_layout.addLayout(btn_layout)

        # Storage for imported entries and checkboxes
        self.psych_imported_entries = []
        self.psych_import_checkboxes = []
        self._psych_current_filter = None
        self._psych_all_categorized = []

        scroll_layout.addWidget(self.psych_import_section)
        scroll_layout.addStretch()
        scroll_area.setWidget(scroll_content)
        main_layout.addWidget(scroll_area, 1)

        return popup

    def _populate_psych_imports(self, raw_notes: list):
        """Populate the psych imports section with filtered notes matching psychiatric history terms."""
        import re
        from datetime import datetime, timedelta

        if not raw_notes:
            print("[MOJ-LEAVE] No raw notes for psych imports")
            self._clear_psych_imports()
            return

        # Psychiatric history categories with keywords and colors
        PSYCH_CATEGORIES = {
            "Admission": [
                "admitted", "admission", "readmitted", "readmission", "informal admission",
                "voluntary admission", "inpatient", "transferred to"
            ],
            "Section": [
                "section 2", "section 3", "section 37", "section 41", "s2", "s3", "s37", "s41",
                "detained under", "mha", "mental health act", "sectioned"
            ],
            "Treatment": [
                "medication", "clozapine", "depot", "antipsychotic", "olanzapine", "risperidone",
                "aripiprazole", "treatment", "ect", "electroconvulsive"
            ],
            "Diagnosis": [
                "diagnosis", "diagnosed", "schizophrenia", "bipolar", "psychosis", "psychotic",
                "schizoaffective", "personality disorder", "depression", "anxiety"
            ],
            "Discharge": [
                "discharged", "discharge", "conditional discharge", "absolute discharge",
                "community treatment order", "cto", "released"
            ],
            "Hospital": [
                "hospital", "unit", "ward", "rampton", "broadmoor", "ashworth",
                "medium secure", "low secure", "high secure", "picu"
            ]
        }

        CATEGORY_COLORS = {
            "Admission": "#7c3aed",    # Purple
            "Section": "#dc2626",      # Red
            "Treatment": "#0891b2",    # Cyan
            "Diagnosis": "#059669",    # Green
            "Discharge": "#d97706",    # Amber
            "Hospital": "#3b82f6"      # Blue
        }

        self._psych_categories = PSYCH_CATEGORIES
        self._psych_category_colors = CATEGORY_COLORS

        def parse_note_date(date_val):
            if isinstance(date_val, datetime):
                return date_val
            if not date_val:
                return None
            date_str = str(date_val).strip()
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"]:
                try:
                    return datetime.strptime(date_str.split()[0] if ' ' in date_str else date_str, fmt.split()[0])
                except:
                    pass
            return None

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def get_matching_categories(text: str) -> list:
            if not text:
                return []
            text_lower = text.lower()
            matched = []
            for cat, keywords in PSYCH_CATEGORIES.items():
                if any(kw in text_lower for kw in keywords):
                    matched.append(cat)
            return matched

        def extract_snippet(text: str, categories: list) -> str:
            if not text:
                return ""
            lines = text.split('\n')
            lines = [l.strip() for l in lines if l.strip()]
            if not lines:
                return text[:300]
            result_lines = lines[:2]
            all_keywords = []
            for cat in categories:
                all_keywords.extend(PSYCH_CATEGORIES.get(cat, []))
            for i, line in enumerate(lines[2:], start=2):
                line_lower = line.lower()
                if any(kw in line_lower for kw in all_keywords):
                    if result_lines and i > len(result_lines) and result_lines[-1] != "...":
                        result_lines.append("...")
                    result_lines.append(line)
            return '\n'.join(result_lines)

        # Filter and categorize notes
        categorized = []
        seen_texts = set()

        for n in raw_notes:
            note_date = parse_note_date(n.get("date") or n.get("datetime"))
            full_text = n.get("body", "") or n.get("text", "") or n.get("content", "")
            if not full_text:
                continue

            categories = get_matching_categories(full_text)
            if not categories:
                continue

            text_sig = full_text[:200].strip()
            if text_sig in seen_texts:
                continue
            seen_texts.add(text_sig)

            snippet = extract_snippet(full_text, categories)
            categorized.append({
                "text": full_text,
                "snippet": snippet,
                "date": n.get("date") or n.get("datetime"),
                "date_obj": note_date,
                "categories": categories
            })

        categorized.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)
        print(f"[MOJ-LEAVE] Found {len(categorized)} psych history entries")

        # Limit to prevent performance issues (show most recent 150)
        MAX_DISPLAY = 150
        if len(categorized) > MAX_DISPLAY:
            print(f"[MOJ-LEAVE] Limiting psych display to {MAX_DISPLAY} most recent entries")
            categorized = categorized[:MAX_DISPLAY]

        self._psych_all_categorized = categorized
        self._display_psych_entries(categorized, format_date_nice, None)

    def _display_psych_entries(self, categorized: list, format_date_nice, filter_category: str = None):
        """Display psych entries with optional category filter."""
        import re
        import html

        if not hasattr(self, 'psych_import_layout') or not self.psych_import_layout:
            return

        CATEGORY_COLORS = self._psych_category_colors
        PSYCH_CATEGORIES = self._psych_categories

        def highlight_keywords(text, matched_categories):
            escaped = html.escape(text)
            keywords_to_highlight = []
            for cat in matched_categories:
                keywords_to_highlight.extend(PSYCH_CATEGORIES.get(cat, []))
            keywords_to_highlight = sorted(set(keywords_to_highlight), key=len, reverse=True)
            for kw in keywords_to_highlight:
                color = CATEGORY_COLORS.get(matched_categories[0], "#fef08a") if matched_categories else "#fef08a"
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m: f'<span style="background-color: {color}; color: white; padding: 1px 2px; border-radius: 2px; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            return escaped.replace('\n', '<br>')

        # Clear existing content
        while self.psych_import_layout.count():
            item = self.psych_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                while item.layout().count():
                    sub_item = item.layout().takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()

        self.psych_import_checkboxes = []

        if filter_category:
            display_entries = [(e, e["categories"]) for e in categorized if filter_category in e.get("categories", [])]
        else:
            display_entries = [(e, e["categories"]) for e in categorized]

        if hasattr(self, 'psych_import_subtitle'):
            self.psych_import_subtitle.setText(f"Showing {len(display_entries)} psych history entries - tick to include in preview")

        if not display_entries:
            placeholder = QLabel("No psychiatric history entries found. Use Import Data to load clinical notes.")
            placeholder.setStyleSheet("color: #92400e; font-style: italic; font-size: 11px; background: transparent;")
            self.psych_import_layout.addWidget(placeholder)
            return

        # Filter header row (red box when filter active)
        if filter_category:
            filter_frame = QFrame()
            filter_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #dc2626; border-radius: 4px; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(8, 4, 8, 4)
            filter_row.setSpacing(8)

            filter_label = QLabel("Filtered by:")
            filter_label.setStyleSheet("font-size: 10px; color: #991b1b; background: transparent;")
            filter_row.addWidget(filter_label)

            color = CATEGORY_COLORS.get(filter_category, "#6b7280")
            cat_label = QLabel(filter_category)
            cat_label.setStyleSheet(f"font-size: 10px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(cat_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setStyleSheet("""
                QPushButton { font-size: 9px; color: #dc2626; background: transparent; border: 1px solid #dc2626; border-radius: 3px; padding: 2px 6px; }
                QPushButton:hover { background: #fee2e2; }
            """)
            remove_btn.clicked.connect(lambda: self._filter_psych_by_category(None))
            filter_row.addWidget(remove_btn)
            self.psych_import_layout.addWidget(filter_frame)

        # Select All / Deselect All buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(6)

        select_all_btn = QPushButton("Select All")
        select_all_btn.setStyleSheet("""
            QPushButton { background: #d97706; color: white; border: none; border-radius: 4px; padding: 4px 10px; font-size: 10px; font-weight: 600; }
            QPushButton:hover { background: #b45309; }
        """)
        select_all_btn.clicked.connect(self._select_all_psych_imports)
        btn_row.addWidget(select_all_btn)

        deselect_all_btn = QPushButton("Deselect All")
        deselect_all_btn.setStyleSheet("""
            QPushButton { background: #92400e; color: white; border: none; border-radius: 4px; padding: 4px 10px; font-size: 10px; font-weight: 600; }
            QPushButton:hover { background: #78350f; }
        """)
        deselect_all_btn.clicked.connect(self._deselect_all_psych_imports)
        btn_row.addWidget(deselect_all_btn)
        btn_row.addStretch()
        self.psych_import_layout.addLayout(btn_row)

        # Add each entry
        for entry, categories in display_entries:
            full_text = entry["text"]
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            display_text = entry.get("snippet", full_text[:150] + "..." if len(full_text) > 150 else full_text)

            entry_frame = QFrame()
            border_color = CATEGORY_COLORS.get(categories[0], "#fcd34d") if categories else "#fcd34d"
            entry_frame.setStyleSheet(f"QFrame {{ background: white; border: 1px solid {border_color}; border-radius: 4px; }}")

            entry_layout_h = QHBoxLayout(entry_frame)
            entry_layout_h.setContentsMargins(6, 4, 6, 4)
            entry_layout_h.setSpacing(6)

            cb = QCheckBox()
            cb.setChecked(False)
            cb.setStyleSheet("QCheckBox { background: transparent; }")
            cb.stateChanged.connect(lambda state: self._update_psych_preview_from_imports())
            entry_layout_h.addWidget(cb)

            content_layout = QVBoxLayout()
            content_layout.setSpacing(2)

            header_layout = QHBoxLayout()
            header_layout.setSpacing(4)

            if categories:
                for cat in categories:
                    tag = QPushButton(cat)
                    tag.setCursor(Qt.CursorShape.PointingHandCursor)
                    color = CATEGORY_COLORS.get(cat, "#6b7280")
                    tag.setStyleSheet(f"""
                        QPushButton {{ font-size: 9px; font-weight: 600; color: white; background: {color}; padding: 1px 4px; border-radius: 3px; border: none; }}
                        QPushButton:hover {{ background: {color}; opacity: 0.8; }}
                    """)
                    tag.clicked.connect(lambda checked, c=cat: self._filter_psych_by_category(c))
                    header_layout.addWidget(tag)

            header_layout.addStretch()

            expand_btn = QPushButton("+")
            expand_btn.setFixedSize(20, 20)
            expand_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            expand_btn.setStyleSheet("""
                QPushButton { font-size: 14px; font-weight: 700; color: #6b7280; background: #f3f4f6; border: 1px solid #d1d5db; border-radius: 4px; }
                QPushButton:hover { background: #e5e7eb; color: #374151; }
            """)
            header_layout.addWidget(expand_btn)
            content_layout.addLayout(header_layout)

            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            if date_str:
                date_lbl = QLabel(date_str)
                date_lbl.setStyleSheet("font-size: 10px; font-weight: 600; color: #92400e; background: transparent;")
                content_layout.addWidget(date_lbl)

            text_lbl = QLabel(display_text[:200] + "..." if len(display_text) > 200 else display_text)
            text_lbl.setWordWrap(True)
            text_lbl.setStyleSheet("font-size: 10px; color: #374151; background: transparent;")
            text_lbl.setTextFormat(Qt.TextFormat.PlainText)
            content_layout.addWidget(text_lbl)

            entry_frame._is_expanded = False
            entry_frame._snippet_text = display_text[:200] + "..." if len(display_text) > 200 else display_text
            entry_frame._full_highlighted = highlight_keywords(full_text, categories)
            entry_frame._text_lbl = text_lbl
            entry_frame._expand_btn = expand_btn

            def make_toggle(ef, tl, eb):
                def toggle_expand():
                    if ef._is_expanded:
                        tl.setTextFormat(Qt.TextFormat.PlainText)
                        tl.setText(ef._snippet_text)
                        eb.setText("+")
                        ef._is_expanded = False
                    else:
                        tl.setTextFormat(Qt.TextFormat.RichText)
                        tl.setText(ef._full_highlighted)
                        eb.setText("−")
                        ef._is_expanded = True
                return toggle_expand

            expand_btn.clicked.connect(make_toggle(entry_frame, text_lbl, expand_btn))
            entry_layout_h.addLayout(content_layout, 1)
            self.psych_import_layout.addWidget(entry_frame)
            self.psych_import_checkboxes.append((cb, entry))

    def _filter_psych_by_category(self, category: str):
        self._psych_current_filter = category
        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")
        self._display_psych_entries(self._psych_all_categorized, format_date_nice, category)

    def _select_all_psych_imports(self):
        for checkbox, _ in self.psych_import_checkboxes:
            checkbox.setChecked(True)

    def _deselect_all_psych_imports(self):
        for checkbox, _ in self.psych_import_checkboxes:
            checkbox.setChecked(False)

    def _clear_psych_imports(self):
        self.psych_imported_entries = []
        self.psych_import_checkboxes = []
        self._psych_all_categorized = []
        self._psych_current_filter = None

        if not hasattr(self, 'psych_import_layout') or not self.psych_import_layout:
            return

        while self.psych_import_layout.count():
            item = self.psych_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
            elif item.layout():
                while item.layout().count():
                    sub_item = item.layout().takeAt(0)
                    if sub_item.widget():
                        sub_item.widget().deleteLater()

        placeholder = QLabel("No entries found. Use Import Data to load clinical notes.")
        placeholder.setStyleSheet("color: #92400e; font-style: italic; font-size: 11px; background: transparent;")
        self.psych_import_layout.addWidget(placeholder)

        if hasattr(self, 'psych_import_subtitle'):
            self.psych_import_subtitle.setText("Imported psychiatric history entries - tick to include in preview")

        if hasattr(self, 'hospital_admissions'):
            self.hospital_admissions.clear()

    def _update_psych_preview_from_imports(self):
        """Auto-update the psych history preview when import checkboxes change."""
        if not hasattr(self, 'psych_import_checkboxes') or not self.psych_import_checkboxes:
            return

        texts = []
        for checkbox, entry in self.psych_import_checkboxes:
            if checkbox.isChecked():
                text = entry.get("text", "").strip()
                date_obj = entry.get("date_obj")
                if text:
                    if date_obj:
                        date_str = date_obj.strftime("%d/%m/%Y")
                        texts.append(f"[{date_str}] {text}")
                    else:
                        texts.append(text)

        if hasattr(self, 'hospital_admissions'):
            if texts:
                self.hospital_admissions.setPlainText("\n\n".join(texts))
            else:
                self.hospital_admissions.clear()

    def _create_index_offence_popup(self) -> QWidget:
        """Create index offence popup with imported forensic history panel."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { border: none; background: white; } QWidget { background: white; }")

        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)

        # === PREVIEW SECTION (fixed height) ===
        preview_container = QFrame()
        preview_container.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; }")
        preview_container.setMaximumHeight(160)
        preview_layout = QVBoxLayout(preview_container)
        preview_layout.setContentsMargins(12, 8, 12, 8)
        preview_layout.setSpacing(4)

        preview_header_row = QHBoxLayout()
        preview_header_row.setContentsMargins(0, 0, 0, 0)

        preview_header = QLabel("Preview")
        preview_header.setStyleSheet("font-size: 13px; font-weight: 600; color: #991b1b;")
        preview_header_row.addWidget(preview_header)
        preview_header_row.addStretch()

        self.index_offence_send_btn = QPushButton("Send to Card")
        self.index_offence_send_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                padding: 6px 14px;
                border: none;
                border-radius: 5px;
                font-weight: 600;
                font-size: 11px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        self.index_offence_send_btn.clicked.connect(lambda: self._send_to_card("index_offence"))
        preview_header_row.addWidget(self.index_offence_send_btn)

        preview_layout.addLayout(preview_header_row)

        self.index_offence_preview = QTextEdit()
        self.index_offence_preview.setReadOnly(True)
        self.index_offence_preview.setMinimumHeight(60)
        self.index_offence_preview.setMaximumHeight(100)
        self.index_offence_preview.setStyleSheet("""
            QTextEdit { background: #1f2937; border: none; border-radius: 4px; padding: 8px; font-size: 12px; color: white; }
        """)
        preview_layout.addWidget(self.index_offence_preview)

        layout.addWidget(preview_container)

        # === INPUT SECTION ===
        layout.addWidget(QLabel("Index Offence and Forensic History:"))

        self.index_offence = QTextEdit()
        self.index_offence.setPlaceholderText("Enter index offence and forensic history...")
        self.index_offence.setMinimumHeight(150)
        self.index_offence.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 12px;
            }
        """)
        zoom_row = create_zoom_row(self.index_offence, base_size=12)
        layout.addLayout(zoom_row)
        layout.addWidget(self.index_offence)

        # Imported data panel (red theme)
        import_panel = QFrame()
        import_panel.setStyleSheet("QFrame { background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 8px; }")
        import_layout = QVBoxLayout(import_panel)
        import_layout.setContentsMargins(10, 8, 10, 8)
        import_layout.setSpacing(6)

        header_lbl = QLabel("Forensic History")
        header_lbl.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        import_layout.addWidget(header_lbl)

        subtitle_lbl = QLabel("Imported data from notes")
        subtitle_lbl.setStyleSheet("font-size: 10px; color: #6b7280; font-style: italic;")
        import_layout.addWidget(subtitle_lbl)

        self.forensic_import_text = QTextEdit()
        self.forensic_import_text.setPlaceholderText("Imported forensic history will appear here after data extraction...")
        self.forensic_import_text.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #e5e7eb; border-radius: 4px; font-size: 10px; padding: 4px; }
        """)
        self.forensic_import_text.setMinimumHeight(80)
        forensic_import_zoom = create_zoom_row(self.forensic_import_text, base_size=10)
        import_layout.addLayout(forensic_import_zoom)
        import_layout.addWidget(self.forensic_import_text)

        # Summary frame
        self.forensic_summary_frame = QFrame()
        self.forensic_summary_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #f59e0b; border-radius: 6px; }")
        self.forensic_summary_frame.hide()
        summary_layout = QVBoxLayout(self.forensic_summary_frame)
        summary_layout.setContentsMargins(6, 6, 6, 6)
        summary_layout.setSpacing(4)

        summary_header = QHBoxLayout()
        summary_label = QLabel("Summary")
        summary_label.setStyleSheet("font-weight: 600; color: #92400e; font-size: 10px;")
        summary_header.addWidget(summary_label)
        summary_header.addStretch()

        copy_forensic_btn = QPushButton("Copy")
        copy_forensic_btn.setFixedWidth(40)
        copy_forensic_btn.setStyleSheet("""
            QPushButton { background: #f59e0b; color: white; border: none; padding: 2px 6px; border-radius: 4px; font-size: 9px; }
            QPushButton:hover { background: #d97706; }
        """)
        copy_forensic_btn.clicked.connect(self._copy_forensic_summary)
        summary_header.addWidget(copy_forensic_btn)

        close_forensic_btn = QPushButton("✕")
        close_forensic_btn.setFixedWidth(20)
        close_forensic_btn.setStyleSheet("""
            QPushButton { background: transparent; color: #92400e; border: none; font-size: 12px; font-weight: bold; }
            QPushButton:hover { color: #78350f; background: rgba(0,0,0,0.1); border-radius: 4px; }
        """)
        close_forensic_btn.clicked.connect(lambda: self.forensic_summary_frame.hide())
        summary_header.addWidget(close_forensic_btn)
        summary_layout.addLayout(summary_header)

        self.forensic_summary_text = QTextEdit()
        self.forensic_summary_text.setMaximumHeight(80)
        self.forensic_summary_text.setStyleSheet("QTextEdit { background: white; border: 1px solid #fcd34d; border-radius: 4px; font-size: 10px; }")
        forensic_summary_zoom = create_zoom_row(self.forensic_summary_text, base_size=10)
        summary_layout.addLayout(forensic_summary_zoom)
        summary_layout.addWidget(self.forensic_summary_text)
        import_layout.addWidget(self.forensic_summary_frame)

        # Button row
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        forensic_summary_btn = QPushButton("Summary")
        forensic_summary_btn.setStyleSheet("""
            QPushButton { background: #f59e0b; color: white; border: none; padding: 6px 12px; border-radius: 4px; font-size: 10px; font-weight: 500; }
            QPushButton:hover { background: #d97706; }
        """)
        forensic_summary_btn.clicked.connect(self._generate_forensic_summary)
        btn_row.addWidget(forensic_summary_btn)

        forensic_send_btn = QPushButton("Send to Report")
        forensic_send_btn.setStyleSheet("""
            QPushButton { background: #8b5cf6; color: white; border: none; padding: 6px 12px; border-radius: 4px; font-size: 10px; font-weight: 600; }
            QPushButton:hover { background: #7c3aed; }
        """)
        forensic_send_btn.clicked.connect(self._send_forensic_to_report)
        btn_row.addWidget(forensic_send_btn)
        btn_row.addStretch()
        import_layout.addLayout(btn_row)

        layout.addWidget(import_panel)

        # Connect preview updates
        self.index_offence.textChanged.connect(self._update_index_offence_preview)

        layout.addStretch()
        popup.setWidget(content)
        return popup

    def _update_index_offence_preview(self):
        """Update the index offence preview with content."""
        if not hasattr(self, 'index_offence_preview'):
            return
        text = self.index_offence.toPlainText().strip() if hasattr(self, 'index_offence') else ""
        self.index_offence_preview.setPlainText(text if text else "(No content yet...)")

    def _create_simple_text_popup(self, label_text: str, placeholder: str, attr_name: str, key: str) -> QWidget:
        """Create a simple popup with preview panel and a text area."""
        popup, layout = self._create_popup_container(key)

        layout.addWidget(QLabel(label_text))

        text_edit = QTextEdit()
        text_edit.setPlaceholderText(placeholder)
        text_edit.setMinimumHeight(150)
        text_edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 12px;
            }
        """)
        setattr(self, attr_name, text_edit)
        zoom_row = create_zoom_row(text_edit, base_size=12)
        layout.addLayout(zoom_row)
        layout.addWidget(text_edit)

        layout.addStretch()

        # Setup preview - simply shows the text content
        def generate_text():
            return text_edit.toPlainText().strip()

        self._setup_popup_preview(key, generate_text)
        self._connect_preview_updates(key, [text_edit])

        return popup

    def _create_mental_disorder_popup(self) -> QWidget:
        return self._create_simple_text_popup("Current Mental Disorder:", "Describe current mental disorder...", "mental_disorder", "mental_disorder")

    def _create_attitude_behaviour_popup(self) -> QWidget:
        return self._create_simple_text_popup("Attitude and Behaviour:", "Describe attitude and behaviour...", "attitude_behaviour", "attitude_behaviour")

    def _create_risk_factors_popup(self) -> QWidget:
        return self._create_simple_text_popup("Risk Factors:", "Describe risk factors...", "risk_factors", "risk_factors")

    def _create_medication_popup(self) -> QWidget:
        return self._create_simple_text_popup("Medication:", "List current medications...", "medication", "medication")

    def _create_psychology_popup(self) -> QWidget:
        return self._create_simple_text_popup("Psychology:", "Describe psychological interventions...", "psychology", "psychology")

    def _create_extremism_popup(self) -> QWidget:
        return self._create_simple_text_popup("Extremism:", "Describe any extremism concerns...", "extremism", "extremism")

    def _create_absconding_popup(self) -> QWidget:
        return self._create_simple_text_popup("Absconding:", "Describe absconding risk...", "absconding", "absconding")

    def _create_mappa_popup(self) -> QWidget:
        return self._create_simple_text_popup("MAPPA:", "Enter MAPPA details...", "mappa_text", "mappa")

    def _create_victims_popup(self) -> QWidget:
        return self._create_simple_text_popup("Victims:", "Enter victim information...", "victims_text", "victims")

    def _create_transferred_prisoners_popup(self) -> QWidget:
        return self._create_simple_text_popup("Transferred Prisoners:", "Enter transferred prisoner details...", "transferred_prisoners_text", "transferred_prisoners")

    def _create_fitness_to_plead_popup(self) -> QWidget:
        return self._create_simple_text_popup("Fitness to Plead:", "Enter fitness to plead information...", "fitness_to_plead_text", "fitness_to_plead")

    def _create_additional_comments_popup(self) -> QWidget:
        return self._create_simple_text_popup("Additional Comments:", "Enter any additional comments...", "additional_comments_text", "additional_comments")

    def _create_signature_popup(self) -> QWidget:
        """Create signature popup with preview panel."""
        key = "signature"
        popup, layout = self._create_popup_container(key)
        field_style = "background: white; padding: 8px; border: 1px solid #d1d5db; border-radius: 6px;"

        layout.addWidget(QLabel("Signature:"))

        # Name
        layout.addWidget(QLabel("Name:"))
        self.sig_name = QLineEdit()
        self.sig_name.setStyleSheet(field_style)
        layout.addWidget(self.sig_name)

        # Role
        layout.addWidget(QLabel("Role:"))
        self.sig_role = QLineEdit()
        self.sig_role.setStyleSheet(field_style)
        layout.addWidget(self.sig_role)

        # Date
        layout.addWidget(QLabel("Date:"))
        self.sig_date = QDateEdit()
        self.sig_date.setCalendarPopup(True)
        self.sig_date.setDate(QDate.currentDate())
        self.sig_date.setDisplayFormat("dd/MM/yyyy")
        self.sig_date.setStyleSheet(field_style)
        layout.addWidget(self.sig_date)

        layout.addStretch()

        # Setup preview
        def generate_signature():
            parts = []
            name = self.sig_name.text().strip()
            role = self.sig_role.text().strip()
            date = self.sig_date.date().toString("dd/MM/yyyy")
            if name:
                parts.append(f"Name: {name}")
            if role:
                parts.append(f"Role: {role}")
            if date:
                parts.append(f"Date: {date}")
            return "\n".join(parts) if parts else ""

        self._setup_popup_preview(key, generate_signature)
        self._connect_preview_updates(key, [self.sig_name, self.sig_role, self.sig_date])

        return popup

    def _create_annex_a_popup(self) -> QWidget:
        return self._create_simple_text_popup("Annex A - Victim Liaison:", "Enter victim liaison details...", "annex_a_text", "annex_a")

    # ================================================================
    # CARD UPDATE METHODS
    # ================================================================

    def _update_card_preview(self, key: str):
        """Update card preview based on popup fields."""
        if key not in self.cards:
            return

        card = self.cards[key]
        text = ""

        if key == "patient_details" and hasattr(self, 'patient_name'):
            name = self.patient_name.text()
            dob = self.patient_dob.date().toString("dd/MM/yyyy") if hasattr(self, 'patient_dob') else ""
            text = f"Name: {name}\nDOB: {dob}"
        elif key == "rc_details" and hasattr(self, 'rc_name'):
            text = f"RC: {self.rc_name.text()}"
        elif key == "leave_type":
            parts = []
            if hasattr(self, 'compassionate_cb') and self.compassionate_cb.isChecked():
                parts.append("Compassionate")
            if hasattr(self, 'escorted_radio') and self.escorted_radio.isChecked():
                parts.append("Escorted")
            elif hasattr(self, 'unescorted_radio') and self.unescorted_radio.isChecked():
                parts.append("Unescorted")
            if hasattr(self, 'leave_type_checkboxes'):
                selected = [cb.text() for cb in self.leave_type_checkboxes.values() if cb.isChecked()]
                parts.extend(selected)
            text = ", ".join(parts) if parts else ""
        elif key == "documents" and hasattr(self, 'documents_checkboxes'):
            selected = [cb.text() for cb in self.documents_checkboxes.values() if cb.isChecked()]
            text = ", ".join(selected) if selected else ""

        if text:
            card.editor.setPlainText(text)

    def _send_to_card(self, key: str):
        """Send popup content to card."""
        if key not in self.cards:
            return

        card = self.cards[key]
        text = ""

        if key == "patient_details":
            parts = []
            if hasattr(self, 'patient_name') and self.patient_name.text():
                parts.append(f"Name: {self.patient_name.text()}")
            if hasattr(self, 'patient_dob'):
                parts.append(f"DOB: {self.patient_dob.date().toString('dd/MM/yyyy')}")
            if hasattr(self, 'hospital_name') and self.hospital_name.text():
                parts.append(f"Hospital: {self.hospital_name.text()}")
            if hasattr(self, 'ward') and self.ward.text():
                parts.append(f"Ward: {self.ward.text()}")
            if hasattr(self, 'mha_section') and self.mha_section.currentText():
                parts.append(f"Section: {self.mha_section.currentText()}")
            if hasattr(self, 'moj_reference') and self.moj_reference.text():
                parts.append(f"MOJ Ref: {self.moj_reference.text()}")
            text = "\n".join(parts)
        elif key == "rc_details":
            parts = []
            if hasattr(self, 'rc_name') and self.rc_name.text():
                parts.append(f"Name: {self.rc_name.text()}")
            if hasattr(self, 'rc_email') and self.rc_email.text():
                parts.append(f"Email: {self.rc_email.text()}")
            if hasattr(self, 'rc_phone') and self.rc_phone.text():
                parts.append(f"Phone: {self.rc_phone.text()}")
            text = "\n".join(parts)
        elif key == "leave_type":
            parts = []
            if hasattr(self, 'compassionate_cb') and self.compassionate_cb.isChecked():
                parts.append("Compassionate/Emergency Leave")
            if hasattr(self, 'escorted_radio') and self.escorted_radio.isChecked():
                escort_status = "Escorted"
            elif hasattr(self, 'unescorted_radio') and self.unescorted_radio.isChecked():
                escort_status = "Unescorted"
            else:
                escort_status = None
            if escort_status and hasattr(self, 'leave_type_checkboxes'):
                selected = [cb.text() for cb in self.leave_type_checkboxes.values() if cb.isChecked()]
                for leave in selected:
                    parts.append(f"{escort_status} {leave}")
            text = "\n".join(parts)
        elif key == "documents":
            parts = []
            if hasattr(self, 'documents_checkboxes'):
                selected = [cb.text() for cb in self.documents_checkboxes.values() if cb.isChecked()]
                parts.extend(selected)
            if hasattr(self, 'other_documents') and self.other_documents.text():
                parts.append(self.other_documents.text())
            text = "\n".join(parts)
        elif key == "purpose" and hasattr(self, 'purpose_preview'):
            text = self.purpose_preview.toPlainText()
        elif key == "overnight" and hasattr(self, 'overnight_preview'):
            text = self.overnight_preview.toPlainText()
        elif key == "escorted_overnight" and hasattr(self, 'escorted_overnight_preview'):
            text = self.escorted_overnight_preview.toPlainText()
        elif key == "compassionate" and hasattr(self, 'compassionate_preview'):
            text = self.compassionate_preview.toPlainText()
        elif key == "leave_report" and hasattr(self, 'leave_report_preview'):
            text = self.leave_report_preview.toPlainText()
        elif key == "procedures" and hasattr(self, 'leave_procedures'):
            text = self.leave_procedures.toPlainText()
        elif key == "hospital_admissions" and hasattr(self, 'hospital_admissions'):
            text = self.hospital_admissions.toPlainText()
        elif key == "index_offence" and hasattr(self, 'index_offence'):
            text = self.index_offence.toPlainText()
        elif key == "mental_disorder" and hasattr(self, 'mental_disorder'):
            text = self.mental_disorder.toPlainText()
        elif key == "attitude_behaviour" and hasattr(self, 'attitude_behaviour'):
            text = self.attitude_behaviour.toPlainText()
        elif key == "risk_factors" and hasattr(self, 'risk_factors'):
            text = self.risk_factors.toPlainText()
        elif key == "medication" and hasattr(self, 'medication'):
            text = self.medication.toPlainText()
        elif key == "psychology" and hasattr(self, 'psychology'):
            text = self.psychology.toPlainText()
        elif key == "extremism" and hasattr(self, 'extremism'):
            text = self.extremism.toPlainText()
        elif key == "absconding" and hasattr(self, 'absconding'):
            text = self.absconding.toPlainText()
        elif key == "signature":
            parts = []
            if hasattr(self, 'sig_name') and self.sig_name.text():
                parts.append(f"Name: {self.sig_name.text()}")
            if hasattr(self, 'sig_role') and self.sig_role.text():
                parts.append(f"Role: {self.sig_role.text()}")
            if hasattr(self, 'sig_date'):
                parts.append(f"Date: {self.sig_date.date().toString('dd/MM/yyyy')}")
            text = "\n".join(parts)
        else:
            # For simple text fields
            attr_mapping = {
                "mappa": "mappa_text",
                "victims": "victims_text",
                "transferred_prisoners": "transferred_prisoners_text",
                "fitness_to_plead": "fitness_to_plead_text",
                "additional_comments": "additional_comments_text",
                "annex_a": "annex_a_text",
            }
            attr_name = attr_mapping.get(key)
            if attr_name and hasattr(self, attr_name):
                text = getattr(self, attr_name).toPlainText()

        card.editor.setPlainText(text)

    # ================================================================
    # SUMMARY AND IMPORT METHODS
    # ================================================================

    def _copy_leave_summary(self):
        from PySide6.QtWidgets import QApplication
        summary = self.leave_summary_text.toPlainText().strip()
        if summary:
            QApplication.clipboard().setText(summary)

    def _generate_leave_summary(self):
        """Generate summary from imported leave data."""
        imported = self.leave_import_text.toPlainText().strip()
        if not imported:
            self.leave_summary_text.setPlainText("No data to summarize.")
            self.leave_summary_frame.show()
            return

        patterns = [
            r'[Ll]eave[^.]{5,150}\.',
            r'[Ee]scort[^.]{5,100}\.',
            r'[Cc]ommunity[^.]{5,150}\.',
            r'[Gg]round[^.]{5,100}\.',
        ]

        found = []
        for pattern in patterns:
            matches = re.findall(pattern, imported)
            for m in matches[:2]:
                clean = " ".join(m.split())
                if clean not in found:
                    found.append(clean)

        if found:
            summary = " ".join(found[:6])
        else:
            sentences = re.findall(r'[^.]+\.', imported)
            summary = " ".join(s.strip() for s in sentences[:4])

        self.leave_summary_text.setPlainText(summary if summary else "No relevant leave information found.")
        self.leave_summary_frame.show()

    def _send_leave_import_to_report(self):
        """Send leave import data to preview."""
        # This method is kept for compatibility but the new structure
        # uses leave_report_preview which auto-generates from input fields
        pass

    def _copy_psych_summary(self):
        from PySide6.QtWidgets import QApplication
        summary = self.psych_summary_text.toPlainText().strip()
        if summary:
            QApplication.clipboard().setText(summary)

    def _generate_psych_summary(self):
        """Generate summary from imported psychiatric history."""
        imported = self.psych_import_text.toPlainText().strip()
        if not imported:
            self.psych_summary_text.setPlainText("No data to summarize.")
            self.psych_summary_frame.show()
            return

        patterns = [
            r'[Aa]dmitted\s+to[^.]{5,150}\.',
            r'[Dd]etained\s+under[^.]{5,100}\.',
            r'[Ii]npatient[^.]{5,150}\.',
            r'[Ss]ection\s*\d+[^.]{5,100}\.',
            r'[Pp]sychiatric\s+(?:admission|hospital|ward)[^.]{5,150}\.',
            r'[Dd]iagnos(?:is|ed)[^.]{5,150}\.',
        ]

        found = []
        for pattern in patterns:
            matches = re.findall(pattern, imported)
            for m in matches[:2]:
                clean = " ".join(m.split())
                if clean not in found:
                    found.append(clean)

        if found:
            summary = " ".join(found[:6])
        else:
            sentences = re.findall(r'[^.]+\.', imported)
            summary = " ".join(s.strip() for s in sentences[:4])

        self.psych_summary_text.setPlainText(summary if summary else "No relevant psychiatric history found.")
        self.psych_summary_frame.show()

    def _send_psych_to_report(self):
        """Send psychiatric data or summary to report."""
        summary = self.psych_summary_text.toPlainText().strip()
        if summary:
            current = self.hospital_admissions.toPlainText()
            self.hospital_admissions.setPlainText(current + "\n\n" + summary if current else summary)

    def _copy_forensic_summary(self):
        from PySide6.QtWidgets import QApplication
        summary = self.forensic_summary_text.toPlainText().strip()
        if summary:
            QApplication.clipboard().setText(summary)

    def _generate_forensic_summary(self):
        """Generate summary from imported forensic history."""
        imported = self.forensic_import_text.toPlainText().strip()
        if not imported:
            self.forensic_summary_text.setPlainText("No data to summarize.")
            self.forensic_summary_frame.show()
            return

        patterns = [
            r'[Ii]ndex\s+offence[^.]{5,200}\.',
            r'[Cc]onvicted\s+(?:of|for)[^.]{5,150}\.',
            r'[Aa]rrested\s+(?:for|on)[^.]{5,150}\.',
            r'[Ss]entenced\s+to[^.]{5,150}\.',
            r'[Pp]rison[^.]{5,100}\.',
            r'[Oo]ffence[^.]{5,150}\.',
        ]

        found = []
        for pattern in patterns:
            matches = re.findall(pattern, imported)
            for m in matches[:2]:
                clean = " ".join(m.split())
                if clean not in found:
                    found.append(clean)

        if found:
            summary = " ".join(found[:6])
        else:
            sentences = re.findall(r'[^.]+\.', imported)
            summary = " ".join(s.strip() for s in sentences[:4])

        self.forensic_summary_text.setPlainText(summary if summary else "No relevant forensic history found.")
        self.forensic_summary_frame.show()

    def _send_forensic_to_report(self):
        """Send forensic data or summary to report."""
        summary = self.forensic_summary_text.toPlainText().strip()
        if summary:
            current = self.index_offence.toPlainText()
            self.index_offence.setPlainText(current + "\n\n" + summary if current else summary)

    # ================================================================
    # DATA EXTRACTOR
    # ================================================================

    def _get_or_create_data_extractor(self):
        """Get or create persistent data extractor instance."""
        if not hasattr(self, '_data_extractor') or self._data_extractor is None:
            try:
                from data_extractor_popup import DataExtractorPopup
                self._data_extractor = DataExtractorPopup(parent=self)
                self._data_extractor.setWindowTitle("Data Extractor - MOJ Leave Application")
                self._data_extractor.setMinimumSize(800, 600)
                # Connect the data extraction signal
                if hasattr(self._data_extractor, 'data_extracted'):
                    self._data_extractor.data_extracted.connect(self._on_data_extracted)
            except ImportError:
                from PySide6.QtWidgets import QMessageBox
                QMessageBox.warning(self, "Import Error", "Data extractor module not available.")
                return None
        return self._data_extractor

    def _import_data(self):
        """Import file - opens file picker directly and processes through data extractor."""
        from PySide6.QtWidgets import QFileDialog, QMessageBox

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Import File",
            "",
            "All Supported Files (*.pdf *.docx *.doc *.txt *.rtf *.xls *.xlsx);;PDF Files (*.pdf);;Word Documents (*.docx *.doc);;Excel Files (*.xls *.xlsx);;All Files (*)"
        )

        if not file_path:
            return

        # Get or create persistent data extractor
        extractor = self._get_or_create_data_extractor()
        if extractor is None:
            return

        try:
            # Load the file directly into the extractor
            if hasattr(extractor, 'load_file'):
                extractor.load_file(file_path)
            elif hasattr(extractor, 'process_file'):
                extractor.process_file(file_path)
            # Show the extractor window
            extractor.show()
            extractor.raise_()
            extractor.activateWindow()
        except Exception as e:
            QMessageBox.warning(self, "Import Error", f"Failed to import file: {str(e)}")

    def _view_data(self):
        """View data extractor popup (persists loaded data)."""
        extractor = self._get_or_create_data_extractor()
        if extractor is None:
            return

        # Show the extractor window (data persists)
        extractor.show()
        extractor.raise_()
        extractor.activateWindow()

    def _on_data_extracted(self, panel_data: dict):
        """Handle extracted data - insert into text fields and populate import sections."""
        categories = panel_data.get("categories", {})
        print(f"[MOJ-LEAVE] Processing {len(categories)} categories")

        # Get raw notes from the data extractor for populating import sections
        raw_notes = []
        if hasattr(self, '_data_extractor') and self._data_extractor:
            raw_notes = getattr(self._data_extractor, 'notes', [])
            print(f"[MOJ-LEAVE] Got {len(raw_notes)} raw notes from extractor")

        # Populate 4a Past Psychiatric History imports section
        if raw_notes:
            self._populate_psych_imports(raw_notes)
            print(f"[MOJ-LEAVE] Populated 4a psych imports with {len(raw_notes)} notes")

        # Map category names to text widgets
        CATEGORY_MAP = {
            "Past Psychiatric History": "hospital_admissions",
            "Medication History": "hospital_admissions",
            "Forensic History": "index_offence",
            "Risk": "leave_report_preview",
            "History of Presenting Complaint": "leave_report_preview",
            "Summary": "leave_report_preview",
            "Plan": "leave_report_preview",
        }

        # Insert extracted data into appropriate text fields
        for cat_key, payload in categories.items():
            if not isinstance(payload, dict):
                continue

            cat_name = payload.get("name", "")
            widget_name = CATEGORY_MAP.get(cat_name)
            if not widget_name:
                continue

            items = payload.get("items", [])
            if not items:
                continue

            # Get the target text widget
            text_widget = getattr(self, widget_name, None)
            if not text_widget:
                continue

            for item in items:
                text = (item.get("text") or "").strip()
                if text:
                    date = item.get("date")
                    if date:
                        if hasattr(date, "strftime"):
                            date_str = date.strftime("%d %b %Y")
                        else:
                            date_str = str(date)
                        text_widget.insertPlainText(f"[{date_str}]\n{text}\n\n")
                    else:
                        text_widget.insertPlainText(text + "\n\n")

            print(f"[MOJ-LEAVE] Inserted {len(items)} items from '{cat_name}'")

    # ================================================================
    # EXPORT / CLEAR
    # ================================================================

    def _export_docx(self):
        """Export form to DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx is not installed.")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export MOJ Leave Form", "", "Word Document (*.docx)"
        )
        if not file_path:
            return

        doc = Document()

        # Title
        title = doc.add_heading("MOJ Leave Application - Restricted Patient", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add each section
        for section_title, key in self.SECTIONS:

            doc.add_heading(section_title, level=2)

            if key in self.cards:
                content = self.cards[key].editor.toPlainText().strip()
                if content:
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph("[No content]")

        doc.save(file_path)
        QMessageBox.information(self, "Export Complete", f"Form exported to:\n{file_path}")

    def _clear_form(self):
        """Clear all form fields."""
        reply = QMessageBox.question(
            self,
            "Clear Form",
            "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply != QMessageBox.StandardButton.Yes:
            return

        # Clear all cards
        for card in self.cards.values():
            card.editor.clear()

        # Clear popup fields
        for attr in dir(self):
            obj = getattr(self, attr, None)
            if isinstance(obj, QTextEdit):
                obj.clear()
            elif isinstance(obj, QLineEdit):
                obj.clear()
            elif isinstance(obj, QCheckBox):
                obj.setChecked(False)
            elif isinstance(obj, QComboBox):
                obj.setCurrentIndex(0)

    # ================================================================
    # STATE MANAGEMENT
    # ================================================================

    def get_state(self) -> dict:
        """Get current form state for saving."""
        state = {"cards": {}}

        for key, card in self.cards.items():
            state["cards"][key] = card.editor.toPlainText()

        # Save popup field values
        state["fields"] = {}
        for attr in ["patient_name", "hospital_number", "hospital_name", "ward", "mha_section", "moj_reference",
                     "rc_name", "rc_email", "rc_phone", "sig_name", "sig_role"]:
            if hasattr(self, attr):
                obj = getattr(self, attr)
                if isinstance(obj, QLineEdit):
                    state["fields"][attr] = obj.text()

        return state

    def set_state(self, state: dict):
        """Restore form state."""
        if "cards" in state:
            for key, text in state["cards"].items():
                if key in self.cards:
                    self.cards[key].editor.setPlainText(text)

        if "fields" in state:
            for attr, value in state["fields"].items():
                if hasattr(self, attr):
                    obj = getattr(self, attr)
                    if isinstance(obj, QLineEdit):
                        obj.setText(value)
