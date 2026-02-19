# ================================================================
#  H1 FORM PAGE — Section 5(2) Report on Hospital In-patient
#  Mental Health Act 1983 - Form H1 Regulation 4(1)(g)
#  Part 1 only - Holding power report
# ================================================================

from __future__ import annotations

from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QGroupBox, QToolButton, QRadioButton,
    QButtonGroup, QComboBox, QCompleter
)

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}


# ================================================================
# TOOLBAR
# ================================================================

class H1Toolbar(QWidget):
    """Toolbar for the H1 Form Page."""

    export_docx = Signal()
    clear_form = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(60)
        self.setStyleSheet("""
            H1Toolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 8, 16, 8)
        layout.setSpacing(12)

        # Export DOCX button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #7c3aed;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #6d28d9; }
            QToolButton:pressed { background: #5b21b6; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Clear Form button
        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #dc2626; }
            QToolButton:pressed { background: #b91c1c; }
        """)
        clear_btn.clicked.connect(self.clear_form.emit)
        layout.addWidget(clear_btn)

        layout.addStretch()


# ================================================================
# MAIN H1 FORM PAGE
# ================================================================

class H1FormPage(QWidget):
    """Page for completing MHA Form H1 - Section 5(2) Report on Hospital In-patient."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()

        self._setup_ui()
        self._prefill_practitioner()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        # Convert sqlite3.Row to dict
        return dict(details)

    def _prefill_practitioner(self):
        """Pre-fill practitioner details from saved settings."""
        if self._my_details:
            self.prac_name.setText(self._my_details.get("name", "") or "")
            self.hospital_name.setText(self._my_details.get("hospital_name", "") or "")
            self.hospital_address.setText(self._my_details.get("hospital_address", "") or "")

    def _setup_ui(self):
        self.setStyleSheet("background: #f3f4f6;")

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(56)
        header.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #1e3a5f, stop:1 #2d5a87);
                border: none;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(16, 0, 16, 0)

        back_btn = QToolButton()
        back_btn.setText("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QToolButton {
                background: rgba(255,255,255,0.15);
                color: white;
                font-size: 14px;
                font-weight: 500;
                border: none;
                border-radius: 4px;
                padding: 8px 12px;
            }
            QToolButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form H1 — Section 5(2) Report on Hospital In-patient")
        title.setStyleSheet("color: white; font-size: 18px; font-weight: 600;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = H1Toolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        self.toolbar.clear_form.connect(self._clear_form)
        main_layout.addWidget(self.toolbar)

        # Scroll area for form content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: #f3f4f6; }")

        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(24, 24, 24, 24)
        content_layout.setSpacing(24)

        # Left column - Main form fields
        left_col = QWidget()
        left_col.setMaximumWidth(500)
        left_layout = QVBoxLayout(left_col)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(16)

        # Part 1 Header
        part1_lbl = QLabel("PART 1")
        part1_lbl.setStyleSheet("font-size: 16px; font-weight: 700; color: #1e3a5f;")
        left_layout.addWidget(part1_lbl)

        part1_desc = QLabel("To be completed by a medical practitioner or approved clinician")
        part1_desc.setStyleSheet("font-size: 12px; color: #6b7280; margin-bottom: 8px;")
        left_layout.addWidget(part1_desc)

        # Patient Section (moved to top)
        patient_frame = QFrame()
        patient_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        patient_layout = QVBoxLayout(patient_frame)
        patient_layout.setContentsMargins(16, 12, 16, 12)
        patient_layout.setSpacing(8)

        patient_title = QLabel("Patient")
        patient_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151;")
        patient_layout.addWidget(patient_title)

        self.patient_name = QLineEdit()
        self.patient_name.setPlaceholderText("Full name of patient")
        self.patient_name.setStyleSheet("padding: 8px; border: 1px solid #d1d5db; border-radius: 4px;")
        patient_layout.addWidget(self.patient_name)

        # Gender selection
        gender_row = QHBoxLayout()
        gender_row.setSpacing(16)

        gender_lbl = QLabel("Gender:")
        gender_lbl.setStyleSheet("font-size: 12px; color: #6b7280;")
        gender_row.addWidget(gender_lbl)

        self.gender_group = QButtonGroup(self)

        self.gender_male = QRadioButton("Male")
        self.gender_male.setStyleSheet("font-size: 12px; color: #374151;")
        self.gender_male.toggled.connect(self._build_narrative)
        self.gender_group.addButton(self.gender_male)
        gender_row.addWidget(self.gender_male)

        self.gender_female = QRadioButton("Female")
        self.gender_female.setStyleSheet("font-size: 12px; color: #374151;")
        self.gender_female.toggled.connect(self._build_narrative)
        self.gender_group.addButton(self.gender_female)
        gender_row.addWidget(self.gender_female)

        self.gender_other = QRadioButton("Other")
        self.gender_other.setStyleSheet("font-size: 12px; color: #374151;")
        self.gender_other.toggled.connect(self._build_narrative)
        self.gender_group.addButton(self.gender_other)
        gender_row.addWidget(self.gender_other)

        gender_row.addStretch()
        patient_layout.addLayout(gender_row)

        left_layout.addWidget(patient_frame)

        # Hospital Section
        hosp_frame = QFrame()
        hosp_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        hosp_layout = QVBoxLayout(hosp_frame)
        hosp_layout.setContentsMargins(16, 12, 16, 12)
        hosp_layout.setSpacing(8)

        hosp_title = QLabel("Hospital")
        hosp_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151;")
        hosp_layout.addWidget(hosp_title)

        self.hospital_name = QLineEdit()
        self.hospital_name.setPlaceholderText("Hospital name")
        self.hospital_name.setStyleSheet("padding: 8px; border: 1px solid #d1d5db; border-radius: 4px;")
        hosp_layout.addWidget(self.hospital_name)

        self.hospital_address = QLineEdit()
        self.hospital_address.setPlaceholderText("Hospital address")
        self.hospital_address.setStyleSheet("padding: 8px; border: 1px solid #d1d5db; border-radius: 4px;")
        hosp_layout.addWidget(self.hospital_address)

        left_layout.addWidget(hosp_frame)

        # Practitioner Section
        prac_frame = QFrame()
        prac_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        prac_layout = QVBoxLayout(prac_frame)
        prac_layout.setContentsMargins(16, 12, 16, 12)
        prac_layout.setSpacing(8)

        prac_title = QLabel("Practitioner")
        prac_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151;")
        prac_layout.addWidget(prac_title)

        self.prac_name = QLineEdit()
        self.prac_name.setPlaceholderText("Full name")
        self.prac_name.setStyleSheet("padding: 8px; border: 1px solid #d1d5db; border-radius: 4px;")
        prac_layout.addWidget(self.prac_name)

        # Practitioner type
        type_lbl = QLabel("I am:")
        type_lbl.setStyleSheet("font-size: 12px; color: #6b7280; margin-top: 4px;")
        prac_layout.addWidget(type_lbl)

        self.prac_type_group = QButtonGroup(self)

        self.prac_in_charge = QRadioButton("(a) In charge of the treatment of this patient")
        self.prac_in_charge.setStyleSheet("font-size: 12px; color: #374151;")
        self.prac_in_charge.setChecked(True)
        self.prac_type_group.addButton(self.prac_in_charge)
        prac_layout.addWidget(self.prac_in_charge)

        self.prac_nominee = QRadioButton("(b) Nominee of the person in charge of treatment")
        self.prac_nominee.setStyleSheet("font-size: 12px; color: #374151;")
        self.prac_type_group.addButton(self.prac_nominee)
        prac_layout.addWidget(self.prac_nominee)

        # Clinician type
        clin_lbl = QLabel("Clinician type:")
        clin_lbl.setStyleSheet("font-size: 12px; color: #6b7280; margin-top: 8px;")
        prac_layout.addWidget(clin_lbl)

        self.clinician_type_group = QButtonGroup(self)

        self.clinician_rmp = QRadioButton("Registered medical practitioner")
        self.clinician_rmp.setStyleSheet("font-size: 12px; color: #374151;")
        self.clinician_rmp.setChecked(True)
        self.clinician_type_group.addButton(self.clinician_rmp)
        prac_layout.addWidget(self.clinician_rmp)

        self.clinician_ac = QRadioButton("Approved clinician (not a registered medical practitioner)")
        self.clinician_ac.setStyleSheet("font-size: 12px; color: #374151;")
        self.clinician_type_group.addButton(self.clinician_ac)
        prac_layout.addWidget(self.clinician_ac)

        left_layout.addWidget(prac_frame)

        left_layout.addStretch()
        content_layout.addWidget(left_col)

        # Right column - Reasons and delivery method
        right_col = QWidget()
        right_layout = QVBoxLayout(right_col)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(16)

        # Reasons Section - horizontal layout with text on left, controls on right
        reasons_frame = QFrame()
        reasons_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        reasons_main_layout = QVBoxLayout(reasons_frame)
        reasons_main_layout.setContentsMargins(16, 12, 16, 12)
        reasons_main_layout.setSpacing(8)

        reasons_title = QLabel("Reasons for Detention")
        reasons_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151;")
        reasons_main_layout.addWidget(reasons_title)

        # Horizontal container for text area (left) and controls (right)
        reasons_h_layout = QHBoxLayout()
        reasons_h_layout.setSpacing(24)

        # Left side - Text area
        self.reasons_text = QTextEdit()
        self.reasons_text.setPlaceholderText("Narrative will build automatically from selections...")
        self.reasons_text.setMinimumHeight(180)
        self.reasons_text.setStyleSheet("padding: 8px; border: 1px solid #d1d5db; border-radius: 4px;")
        reasons_h_layout.addWidget(self.reasons_text, stretch=1)

        # Right side - Controls panel
        controls_widget = QWidget()
        controls_layout = QVBoxLayout(controls_widget)
        controls_layout.setContentsMargins(0, 0, 0, 0)
        controls_layout.setSpacing(8)

        # Diagnosis dropdown
        dx_lbl = QLabel("Diagnosis (ICD-10)")
        dx_lbl.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        controls_layout.addWidget(dx_lbl)

        self.dx_combo = QComboBox()
        self.dx_combo.setEditable(True)
        self.dx_combo.setMinimumWidth(350)
        self.dx_combo.setStyleSheet("padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 11px;")
        self.dx_combo.addItem("", None)
        completer_items = []
        for key, val in sorted(ICD10_DICT.items()):
            # Handle both dict values and string values
            if isinstance(val, dict):
                name = val.get("diagnosis", val.get("name", str(val)))
                code = val.get("icd10", val.get("code", key))
            else:
                name = str(val)
                code = key
            display = f"{code} - {name}"
            self.dx_combo.addItem(display, {"code": code, "name": name})
            completer_items.append(display)
        self.dx_combo.setCurrentIndex(0)
        completer = QCompleter(completer_items)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_combo.setCompleter(completer)
        self.dx_combo.currentIndexChanged.connect(self._build_narrative)
        controls_layout.addWidget(self.dx_combo)

        # Reasons checkboxes
        reasons_lbl = QLabel("Reasons")
        reasons_lbl.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151; margin-top: 8px;")
        controls_layout.addWidget(reasons_lbl)

        reasons_row = QHBoxLayout()
        reasons_row.setSpacing(16)

        self.cb_refusing = QCheckBox("Refusing to remain")
        self.cb_refusing.setStyleSheet("font-size: 11px; color: #374151;")
        self.cb_refusing.stateChanged.connect(self._build_narrative)
        reasons_row.addWidget(self.cb_refusing)

        self.cb_very_unwell = QCheckBox("Very unwell")
        self.cb_very_unwell.setStyleSheet("font-size: 11px; color: #374151;")
        self.cb_very_unwell.stateChanged.connect(self._build_narrative)
        reasons_row.addWidget(self.cb_very_unwell)

        self.cb_acute = QCheckBox("Acute deterioration")
        self.cb_acute.setStyleSheet("font-size: 11px; color: #374151;")
        self.cb_acute.stateChanged.connect(self._build_narrative)
        reasons_row.addWidget(self.cb_acute)

        reasons_row.addStretch()
        controls_layout.addLayout(reasons_row)

        # Significant risk section
        risk_row = QHBoxLayout()
        risk_row.setSpacing(16)

        risk_lbl = QLabel("Significant risk to:")
        risk_lbl.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        risk_row.addWidget(risk_lbl)

        self.cb_risk_self = QCheckBox("Self")
        self.cb_risk_self.setStyleSheet("font-size: 11px; color: #374151;")
        self.cb_risk_self.stateChanged.connect(self._build_narrative)
        risk_row.addWidget(self.cb_risk_self)

        self.cb_risk_others = QCheckBox("Others")
        self.cb_risk_others.setStyleSheet("font-size: 11px; color: #374151;")
        self.cb_risk_others.stateChanged.connect(self._build_narrative)
        risk_row.addWidget(self.cb_risk_others)

        risk_row.addStretch()
        controls_layout.addLayout(risk_row)

        controls_layout.addStretch()
        reasons_h_layout.addWidget(controls_widget, stretch=1)

        reasons_main_layout.addLayout(reasons_h_layout)
        right_layout.addWidget(reasons_frame)

        # Delivery Method Section
        delivery_frame = QFrame()
        delivery_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        delivery_layout = QVBoxLayout(delivery_frame)
        delivery_layout.setContentsMargins(16, 12, 16, 12)
        delivery_layout.setSpacing(8)

        delivery_title = QLabel("Report Furnishing Method")
        delivery_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151;")
        delivery_layout.addWidget(delivery_title)

        self.delivery_group = QButtonGroup(self)

        # Option a - internal mail
        self.delivery_internal = QRadioButton("(a) Internal mail system")
        self.delivery_internal.setStyleSheet("font-size: 12px; color: #374151;")
        self.delivery_internal.setChecked(True)
        self.delivery_internal.toggled.connect(self._on_delivery_changed)
        self.delivery_group.addButton(self.delivery_internal)
        delivery_layout.addWidget(self.delivery_internal)

        # Time for internal mail
        time_row = QHBoxLayout()
        time_row.setContentsMargins(20, 0, 0, 0)
        time_lbl = QLabel("Time consigned:")
        time_lbl.setStyleSheet("font-size: 12px; color: #6b7280;")
        time_row.addWidget(time_lbl)
        self.internal_time = QTimeEdit()
        self.internal_time.setTime(QTime.currentTime())
        self.internal_time.setDisplayFormat("HH:mm")
        self.internal_time.setStyleSheet("padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        time_row.addWidget(self.internal_time)
        time_row.addStretch()
        delivery_layout.addLayout(time_row)

        # Option b - electronic
        self.delivery_electronic = QRadioButton("(b) Electronic communication")
        self.delivery_electronic.setStyleSheet("font-size: 12px; color: #374151;")
        self.delivery_group.addButton(self.delivery_electronic)
        delivery_layout.addWidget(self.delivery_electronic)

        # Option c - hand delivery
        self.delivery_hand = QRadioButton("(c) Delivered by hand")
        self.delivery_hand.setStyleSheet("font-size: 12px; color: #374151;")
        self.delivery_group.addButton(self.delivery_hand)
        delivery_layout.addWidget(self.delivery_hand)

        right_layout.addWidget(delivery_frame)

        # Signature Section
        sig_frame = QFrame()
        sig_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        sig_layout = QVBoxLayout(sig_frame)
        sig_layout.setContentsMargins(16, 12, 16, 12)
        sig_layout.setSpacing(8)

        sig_title = QLabel("Signature")
        sig_title.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151;")
        sig_layout.addWidget(sig_title)

        date_row = QHBoxLayout()
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 12px; color: #6b7280;")
        date_row.addWidget(date_lbl)
        self.sig_date = QDateEdit()
        self.sig_date.setDate(QDate.currentDate())
        self.sig_date.setCalendarPopup(True)
        self.sig_date.setDisplayFormat("dd/MM/yyyy")
        self.sig_date.setStyleSheet("padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        date_row.addWidget(self.sig_date)
        date_row.addStretch()
        sig_layout.addLayout(date_row)

        right_layout.addWidget(sig_frame)

        right_layout.addStretch()
        content_layout.addWidget(right_col, stretch=1)

        scroll.setWidget(content)
        main_layout.addWidget(scroll)

    def _on_delivery_changed(self):
        """Enable/disable time field based on delivery method."""
        self.internal_time.setEnabled(self.delivery_internal.isChecked())

    def _build_narrative(self):
        """Build narrative text from diagnosis and tick boxes."""
        parts = []

        # Determine pronouns based on gender
        if self.gender_male.isChecked():
            subj, poss, reflex = "He", "His", "himself"
        elif self.gender_female.isChecked():
            subj, poss, reflex = "She", "Her", "herself"
        elif self.gender_other.isChecked():
            subj, poss, reflex = "They", "Their", "themselves"
        else:
            # No gender selected - default to they/their
            subj, poss, reflex = "They", "Their", "themselves"

        # Verb conjugation for they vs he/she
        is_verb = "is" if subj != "They" else "are"

        # Get diagnosis
        dx_data = self.dx_combo.currentData()
        if dx_data:
            parts.append(f"The patient suffers from {dx_data['name']} ({dx_data['code']}).")

        # Refusing to remain
        if self.cb_refusing.isChecked():
            parts.append(f"{subj} {is_verb} refusing to remain in hospital informally.")

        # Very unwell and/or acute deterioration
        if self.cb_very_unwell.isChecked() and self.cb_acute.isChecked():
            parts.append(f"{subj} {is_verb} very unwell and suffering an acute deterioration of {poss.lower()} mental state.")
        elif self.cb_very_unwell.isChecked():
            parts.append(f"{subj} {is_verb} currently very unwell.")
        elif self.cb_acute.isChecked():
            parts.append(f"{subj} {is_verb} suffering an acute deterioration of {poss.lower()} mental state.")

        # Risk sentences
        if self.cb_risk_self.isChecked() and self.cb_risk_others.isChecked():
            parts.append(f"{poss} risk to {reflex} and others is significant warranting a mental health act assessment.")
        elif self.cb_risk_self.isChecked():
            parts.append(f"{poss} risk to {reflex} is significant warranting a mental health act assessment.")
        elif self.cb_risk_others.isChecked():
            parts.append(f"{poss} risk to others is significant warranting a mental health act assessment.")

        self.reasons_text.setPlainText(" ".join(parts))

    def _clear_form(self):
        """Clear all form fields."""
        reply = QMessageBox.question(
            self, "Clear Form",
            "Are you sure you want to clear all fields?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.patient_name.clear()
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.dx_combo.setCurrentIndex(0)
            self.cb_refusing.setChecked(False)
            self.cb_very_unwell.setChecked(False)
            self.cb_risk_self.setChecked(False)
            self.cb_risk_others.setChecked(False)
            self.cb_acute.setChecked(False)
            self.reasons_text.clear()
            self.prac_in_charge.setChecked(True)
            self.clinician_rmp.setChecked(True)
            self.delivery_internal.setChecked(True)
            self.internal_time.setTime(QTime.currentTime())
            self.sig_date.setDate(QDate.currentDate())
            # Re-prefill practitioner details
            self._prefill_practitioner()

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form H1",
            f"Form_H1_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_H1_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form H1 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    run = para.add_run(new_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Data goes into WHITESPACE paragraphs below instruction paragraphs
            # Para 5: Hospital data (below para 4 instruction)
            # Para 7: Practitioner name (below para 6 instruction)
            # Para 12: Patient name (below para 11 instruction)
            # Para 15-17: Reasons text (whitespace area)
            # Para 20: Internal mail time
            # Para 23: Signature

            # Hospital - goes in para 5
            hospital_text = self.hospital_name.text()
            if self.hospital_address.text():
                hospital_text += ", " + self.hospital_address.text()
            if hospital_text.strip():
                set_para_text(paragraphs[5], hospital_text)
                highlight_yellow(paragraphs[5])

            # Practitioner name - goes in para 7
            if self.prac_name.text().strip():
                set_para_text(paragraphs[7], self.prac_name.text())
                highlight_yellow(paragraphs[7])

            # Practitioner type - strikethrough non-applicable option
            # Para 9: "the registered medical practitioner/the approved clinician..." - in charge
            # Para 10: "a registered medical practitioner/an approved clinician..." - nominee
            if self.prac_in_charge.isChecked():
                strikethrough_para(paragraphs[10])
            else:
                strikethrough_para(paragraphs[9])

            # Clinician type - strikethrough within the selected paragraph
            # This is handled by the delete phrase instruction in the template

            # Patient name - goes in para 12
            if self.patient_name.text().strip():
                set_para_text(paragraphs[12], self.patient_name.text())
                highlight_yellow(paragraphs[12])

            # Reasons - goes in para 15 (narrative built from tick boxes and diagnosis)
            reasons_text = self.reasons_text.toPlainText().strip()
            if reasons_text:
                set_para_text(paragraphs[15], reasons_text)
                highlight_yellow(paragraphs[15])

            # Delivery method - strikethrough non-applicable options
            # Para 20: internal mail with [time]
            # Para 21: electronic
            # Para 22: hand delivery
            if self.delivery_internal.isChecked():
                strikethrough_para(paragraphs[21])
                strikethrough_para(paragraphs[22])
                # Add time to para 20
                time_str = self.internal_time.time().toString("HH:mm")
                full_text = paragraphs[20].text
                if '[time]' in full_text:
                    new_text = full_text.replace('[time]', time_str)
                    set_para_text(paragraphs[20], new_text)
            elif self.delivery_electronic.isChecked():
                strikethrough_para(paragraphs[20])
                strikethrough_para(paragraphs[22])
            else:  # hand delivery
                strikethrough_para(paragraphs[20])
                strikethrough_para(paragraphs[21])

            # Signature date - para 23
            sig_date = self.sig_date.date().toString("dd MMMM yyyy")
            for run in paragraphs[23].runs:
                run.text = ""
            paragraphs[23].add_run(f"Signed                                                            Date {sig_date}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form H1 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        """Get current form state for saving."""
        gender = "male" if self.gender_male.isChecked() else ("female" if self.gender_female.isChecked() else ("other" if self.gender_other.isChecked() else ""))
        return {
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "prac_name": self.prac_name.text(),
            "prac_in_charge": self.prac_in_charge.isChecked(),
            "clinician_rmp": self.clinician_rmp.isChecked(),
            "patient_name": self.patient_name.text(),
            "gender": gender,
            "reasons": self.reasons_text.toPlainText(),
            "delivery_internal": self.delivery_internal.isChecked(),
            "delivery_electronic": self.delivery_electronic.isChecked(),
            "delivery_hand": self.delivery_hand.isChecked(),
            "internal_time": self.internal_time.time().toString("HH:mm"),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
        }

    def set_state(self, state: dict):
        """Restore form state from saved data."""
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        self.prac_name.setText(state.get("prac_name", ""))

        if state.get("prac_in_charge", True):
            self.prac_in_charge.setChecked(True)
        else:
            self.prac_nominee.setChecked(True)

        if state.get("clinician_rmp", True):
            self.clinician_rmp.setChecked(True)
        else:
            self.clinician_ac.setChecked(True)

        self.patient_name.setText(state.get("patient_name", ""))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        else:
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
        self.reasons_text.setPlainText(state.get("reasons", ""))

        if state.get("delivery_internal", True):
            self.delivery_internal.setChecked(True)
        elif state.get("delivery_electronic", False):
            self.delivery_electronic.setChecked(True)
        else:
            self.delivery_hand.setChecked(True)

        if state.get("internal_time"):
            self.internal_time.setTime(QTime.fromString(state["internal_time"], "HH:mm"))

        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))
