# ================================================================
#  CTO3 FORM PAGE — Notice of Recall to Hospital
#  Mental Health Act 1983 - Form CTO3 Regulation 6(3)(a)
#  Section 17E — Community treatment order: notice of recall
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QToolButton, QRadioButton, QButtonGroup,
    QComboBox, QSpinBox, QCompleter, QStyleFactory, QSlider
)

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


class CTO3FormPage(QWidget):
    """Page for completing MHA Form CTO3 - Notice of Recall to Hospital."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean",
        "Asian",
        "Caucasian",
        "Middle Eastern",
        "Mixed Race",
        "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #dc2626; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO3 — Notice of Recall to Hospital")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        toolbar = QWidget()
        toolbar.setFixedHeight(60)
        toolbar.setStyleSheet("background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12);")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 8, 16, 8)
        tb_layout.setSpacing(12)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #dc2626;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #b91c1c; }
        """)
        export_btn.clicked.connect(self._export_docx)
        tb_layout.addWidget(export_btn)

        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #6b7280;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #4b5563; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        tb_layout.addWidget(clear_btn)
        tb_layout.addStretch()

        main_layout.addWidget(toolbar)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        self._create_section_1_demographics()
        self._create_section_2_patient_hospital()
        self._create_section_3_reason_for_recall()
        self._create_section_4_rc_signature()

        self.form_layout.addStretch()
        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, number: int, title: str, color: str = "#dc2626") -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)

        header_layout = QHBoxLayout()
        header_layout.setSpacing(12)

        number_badge = QLabel(str(number))
        number_badge.setFixedSize(32, 32)
        number_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        number_badge.setStyleSheet(f"""
            background: {color};
            color: white;
            font-size: 14px;
            font-weight: 700;
            border-radius: 16px;
        """)
        header_layout.addWidget(number_badge)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1f2937;")
        header_layout.addWidget(title_lbl)
        header_layout.addStretch()

        layout.addLayout(header_layout)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus { border-color: #dc2626; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 100) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMinimumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #dc2626; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return date_edit

    def _create_time_edit(self) -> QTimeEdit:
        time_edit = QTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("QTimeEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return time_edit

    # ----------------------------------------------------------------
    # SECTION 1: Patient Details
    # ----------------------------------------------------------------
    def _create_section_1_demographics(self):
        frame = self._create_section_frame(1, "Patient Details")
        layout = frame.layout()

        # Patient name and address row
        name_row = QHBoxLayout()
        name_row.setSpacing(12)
        self.patient_name = self._create_line_edit("Patient full name")
        name_row.addWidget(self.patient_name, 1)
        self.patient_address = self._create_line_edit("Patient address")
        name_row.addWidget(self.patient_address, 2)
        layout.addLayout(name_row)

        # Demographics row
        demo_row = QHBoxLayout()
        demo_row.setSpacing(16)

        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        demo_row.addWidget(age_lbl)

        self.age_spin = QSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setFixedWidth(70)
        self.age_spin.setStyleSheet("QSpinBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 13px; }")
        demo_row.addWidget(self.age_spin)

        demo_row.addSpacing(20)

        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("Male")
        self.gender_female = QRadioButton("Female")
        self.gender_other = QRadioButton("Other")
        self.gender_group.addButton(self.gender_male, 0)
        self.gender_group.addButton(self.gender_female, 1)
        self.gender_group.addButton(self.gender_other, 2)

        demo_row.addWidget(self.gender_male)
        demo_row.addWidget(self.gender_female)
        demo_row.addWidget(self.gender_other)

        demo_row.addSpacing(20)

        self.ethnicity_combo = QComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(200)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 13px; }")
        demo_row.addWidget(self.ethnicity_combo)

        demo_row.addStretch()
        layout.addLayout(demo_row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 2: Hospital Details
    # ----------------------------------------------------------------
    def _create_section_2_patient_hospital(self):
        frame = self._create_section_frame(2, "Hospital Details")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(12)
        self.hospital_name = self._create_line_edit("Hospital name")
        row.addWidget(self.hospital_name, 1)
        self.hospital_address = self._create_line_edit("Hospital address")
        row.addWidget(self.hospital_address, 2)
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 3: Reason for Recall (with clinical controls)
    # ----------------------------------------------------------------
    def _create_section_3_reason_for_recall(self):
        frame = self._create_section_frame(3, "Reason for Recall")
        layout = frame.layout()

        split_layout = QHBoxLayout()
        split_layout.setSpacing(16)

        # === LEFT: Grounds Text Area ===
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(6)

        info = QLabel("Click options on the right to auto-generate grounds:")
        info.setStyleSheet("font-size: 10px; color: #6b7280; padding: 4px; background: #fef2f2; border-radius: 4px;")
        left_layout.addWidget(info)

        self.grounds = QTextEdit()
        self.grounds.setPlaceholderText("Grounds for recall will be generated here...")
        self.grounds.setMinimumHeight(350)
        self.grounds.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px; font-size: 12px; }
            QTextEdit:focus { border-color: #dc2626; }
        """)
        left_layout.addWidget(self.grounds)

        split_layout.addWidget(left_container, 3)

        # === RIGHT: Controls Panel ===
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        right_scroll.setFixedWidth(400)

        right_container = QWidget()
        right_container.setFixedWidth(380)
        right_container.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 8, 0)
        right_layout.setSpacing(10)

        # --- Mental Disorder (ICD-10) ---
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(10, 8, 10, 8)
        md_layout.setSpacing(6)

        md_header = QLabel("Mental Disorder")
        md_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        self.dx_boxes = []
        for i in range(2):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText("Primary diagnosis..." if i == 0 else "Secondary (optional)...")

            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(ICD10_DICT.items(), key=lambda x: x[0].lower()):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})

            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(12)
            combo.setStyleSheet("""
                QComboBox { padding: 5px; font-size: 11px; border: 1px solid #d1d5db; border-radius: 4px; background: white; }
                QComboBox QAbstractItemView { min-width: 300px; }
            """)
            combo.currentIndexChanged.connect(self._update_grounds_text)
            md_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        right_layout.addWidget(md_frame)

        # --- Legal Criteria ---
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(10, 8, 10, 8)
        lc_layout.setSpacing(4)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature with sub-options
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._update_grounds_text)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._update_grounds_text)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._update_grounds_text)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree with slider
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(100)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 11px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 11px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_grounds_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        right_layout.addWidget(lc_frame)

        # --- Recall Reason ---
        rr_frame = QFrame()
        rr_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; }")
        rr_layout = QVBoxLayout(rr_frame)
        rr_layout.setContentsMargins(10, 8, 10, 8)
        rr_layout.setSpacing(4)

        rr_header = QLabel("Reason for Recall")
        rr_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #991b1b;")
        rr_layout.addWidget(rr_header)

        # (a) Treatment required AND risk - now a checkbox
        self.reason_a = QCheckBox("(a) Treatment required AND risk")
        self.reason_a.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.reason_a.toggled.connect(self._on_reason_a_toggled)
        rr_layout.addWidget(self.reason_a)

        # Sub-options for reason a
        self.reason_a_options = QWidget()
        reason_a_layout = QVBoxLayout(self.reason_a_options)
        reason_a_layout.setContentsMargins(16, 2, 0, 2)
        reason_a_layout.setSpacing(2)

        self.risk_health_cb = QCheckBox("Risk to mental health")
        self.risk_health_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.risk_health_cb.toggled.connect(self._update_grounds_text)
        reason_a_layout.addWidget(self.risk_health_cb)

        self.risk_safety_cb = QCheckBox("Risk to safety")
        self.risk_safety_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.risk_safety_cb.toggled.connect(self._update_grounds_text)
        reason_a_layout.addWidget(self.risk_safety_cb)

        self.risk_others_cb = QCheckBox("Risk to others")
        self.risk_others_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.risk_others_cb.toggled.connect(self._update_grounds_text)
        reason_a_layout.addWidget(self.risk_others_cb)

        self.reason_a_options.hide()
        rr_layout.addWidget(self.reason_a_options)

        # (b) Failed to comply with condition - now a checkbox
        self.reason_b = QCheckBox("(b) Failed to comply with condition")
        self.reason_b.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.reason_b.toggled.connect(self._on_reason_b_toggled)
        rr_layout.addWidget(self.reason_b)

        # Sub-options for reason b - condition types
        self.reason_b_options = QWidget()
        reason_b_layout = QVBoxLayout(self.reason_b_options)
        reason_b_layout.setContentsMargins(16, 2, 0, 2)
        reason_b_layout.setSpacing(2)

        self.condition_cmht_cb = QCheckBox("Seeing CMHT")
        self.condition_cmht_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.condition_cmht_cb.toggled.connect(self._update_grounds_text)
        reason_b_layout.addWidget(self.condition_cmht_cb)

        self.condition_medication_cb = QCheckBox("Medication")
        self.condition_medication_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.condition_medication_cb.toggled.connect(self._update_grounds_text)
        reason_b_layout.addWidget(self.condition_medication_cb)

        self.condition_residence_cb = QCheckBox("Residence")
        self.condition_residence_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.condition_residence_cb.toggled.connect(self._update_grounds_text)
        reason_b_layout.addWidget(self.condition_residence_cb)

        self.reason_b_options.hide()
        rr_layout.addWidget(self.reason_b_options)

        right_layout.addWidget(rr_frame)
        right_layout.addStretch()

        right_scroll.setWidget(right_container)
        split_layout.addWidget(right_scroll)

        layout.addLayout(split_layout)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 4: RC Signature
    # ----------------------------------------------------------------
    def _create_section_4_rc_signature(self):
        frame = self._create_section_frame(4, "Responsible Clinician")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(12)
        self.rc_name = self._create_line_edit("RC full name")
        row.addWidget(self.rc_name, 1)
        layout.addLayout(row)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(20)
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(date_lbl)
        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        sig_row.addWidget(self.sig_date)

        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(time_lbl)
        self.sig_time = self._create_time_edit()
        self.sig_time.setFixedWidth(100)
        sig_row.addWidget(self.sig_time)
        sig_row.addStretch()
        layout.addLayout(sig_row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # Toggle handlers
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_grounds_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_grounds_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_grounds_text()

    def _on_reason_a_toggled(self, checked):
        self.reason_a_options.setVisible(checked)
        if not checked:
            self.risk_health_cb.setChecked(False)
            self.risk_safety_cb.setChecked(False)
            self.risk_others_cb.setChecked(False)
        self._update_grounds_text()

    def _on_reason_b_toggled(self, checked):
        self.reason_b_options.setVisible(checked)
        if not checked:
            self.condition_cmht_cb.setChecked(False)
            self.condition_medication_cb.setChecked(False)
            self.condition_residence_cb.setChecked(False)
        self._update_grounds_text()

    # ----------------------------------------------------------------
    # Text generation
    # ----------------------------------------------------------------
    def _update_grounds_text(self):
        self.grounds.setPlainText(self._generate_grounds_text())

    def _generate_grounds_text(self) -> str:
        # CTO3 is addressed directly to the patient using "you"
        paragraphs = []

        # Para 1: Demographics + Diagnosis
        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", "").replace("Mixed ", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    diagnoses.append(f"{dx} ({icd})" if icd else dx)

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"You are a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"You suffer from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for you to receive medical treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for you to receive medical treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for you to receive medical treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        # Para 2: Reason for Recall - Treatment required
        para2_parts = []
        if self.reason_a.isChecked():
            para2_parts.append("I am recalling you because you require treatment for your mental disorder in hospital.")
            risk_items = []
            if self.risk_health_cb.isChecked():
                risk_items.append("your mental health")
            if self.risk_safety_cb.isChecked():
                risk_items.append("your own safety")
            if self.risk_others_cb.isChecked():
                risk_items.append("others")
            if risk_items:
                if len(risk_items) == 1:
                    para2_parts.append(f"If not recalled, there would be risks to {risk_items[0]}.")
                elif len(risk_items) == 2:
                    para2_parts.append(f"If not recalled, there would be risks to {risk_items[0]} and {risk_items[1]}.")
                else:
                    para2_parts.append(f"If not recalled, there would be risks to {risk_items[0]}, {risk_items[1]}, and to {risk_items[2]}.")

        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        # Para 3: Failed to comply with condition
        para3_parts = []
        if self.reason_b.isChecked():
            conditions = []
            if self.condition_cmht_cb.isChecked():
                conditions.append("the condition specifying seeing your care coordinator")
            if self.condition_medication_cb.isChecked():
                conditions.append("the CTO condition of taking your medication")
            if self.condition_residence_cb.isChecked():
                conditions.append("your condition of residence as part of the CTO")

            if conditions:
                # Determine prefix based on whether treatment required is also checked
                if self.reason_a.isChecked():
                    prefix = "You have also failed to comply with"
                else:
                    prefix = "You have failed to comply with"

                if len(conditions) == 1:
                    para3_parts.append(f"{prefix} {conditions[0]}.")
                elif len(conditions) == 2:
                    para3_parts.append(f"{prefix} {conditions[0]} and {conditions[1]}.")
                else:
                    para3_parts.append(f"{prefix} {conditions[0]}, {conditions[1]}, and {conditions[2]}.")

        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        return "\n\n".join(paragraphs)

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.patient_name.clear()
            self.hospital_name.clear()
            self.hospital_address.clear()
            for combo in self.dx_boxes:
                combo.setCurrentIndex(0)
            self.patient_address.clear()
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.reason_a.setChecked(False)
            self.reason_b.setChecked(False)
            self.condition_cmht_cb.setChecked(False)
            self.condition_medication_cb.setChecked(False)
            self.condition_residence_cb.setChecked(False)
            self.grounds.clear()
            self.rc_name.clear()
            self.sig_date.setDate(QDate.currentDate())
            self.sig_time.setTime(QTime.currentTime())
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form CTO3",
            f"Form_CTO3_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_CTO3_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO3 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Patient name (para 4)
            if self.patient_name.text().strip():
                set_para_text(paragraphs[4], self.patient_name.text())
                highlight_yellow(paragraphs[4])

            # Hospital (para 6)
            hospital_text = self.hospital_name.text()
            if self.hospital_address.text():
                hospital_text += ", " + self.hospital_address.text()
            if hospital_text.strip():
                set_para_text(paragraphs[6], hospital_text)
                highlight_yellow(paragraphs[6])

            # Strikethrough non-selected reason
            if self.reason_a.isChecked() and not self.reason_b.isChecked():
                strikethrough_para(paragraphs[19])  # Strike out (b)
            elif self.reason_b.isChecked() and not self.reason_a.isChecked():
                strikethrough_para(paragraphs[10])  # Strike out (a) section
                strikethrough_para(paragraphs[11])
                strikethrough_para(paragraphs[12])

            # Grounds (para 14-16)
            if self.grounds.toPlainText().strip():
                set_para_text(paragraphs[14], self.grounds.toPlainText())
                highlight_yellow(paragraphs[14])

            # RC name (para 24)
            if self.rc_name.text().strip():
                set_para_text(paragraphs[24], self.rc_name.text())
                highlight_yellow(paragraphs[24])

            # Date and time (para 25, 26)
            set_para_text(paragraphs[25], self.sig_date.date().toString("dd MMMM yyyy"))
            set_para_text(paragraphs[26], self.sig_time.time().toString("HH:mm"))

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO3 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    def get_state(self) -> dict:
        return {
            "age": self.age_spin.value(),
            "gender": "male" if self.gender_male.isChecked() else "female" if self.gender_female.isChecked() else "other" if self.gender_other.isChecked() else "",
            "ethnicity": self.ethnicity_combo.currentText(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "diagnoses": [combo.currentText() for combo in self.dx_boxes],
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "reason_a": self.reason_a.isChecked(),
            "reason_b": self.reason_b.isChecked(),
            "risk_health": self.risk_health_cb.isChecked(),
            "risk_safety": self.risk_safety_cb.isChecked(),
            "risk_others": self.risk_others_cb.isChecked(),
            "condition_cmht": self.condition_cmht_cb.isChecked(),
            "condition_medication": self.condition_medication_cb.isChecked(),
            "condition_residence": self.condition_residence_cb.isChecked(),
            "grounds": self.grounds.toPlainText(),
            "rc_name": self.rc_name.text(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
            "sig_time": self.sig_time.time().toString("HH:mm"),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.age_spin.setValue(state.get("age", 0))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        ethnicity = state.get("ethnicity", "Ethnicity")
        idx = self.ethnicity_combo.findText(ethnicity)
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.reason_a.setChecked(state.get("reason_a", False))
        self.reason_b.setChecked(state.get("reason_b", False))
        self.risk_health_cb.setChecked(state.get("risk_health", False))
        self.risk_safety_cb.setChecked(state.get("risk_safety", False))
        self.risk_others_cb.setChecked(state.get("risk_others", False))
        self.condition_cmht_cb.setChecked(state.get("condition_cmht", False))
        self.condition_medication_cb.setChecked(state.get("condition_medication", False))
        self.condition_residence_cb.setChecked(state.get("condition_residence", False))
        self.grounds.setPlainText(state.get("grounds", ""))
        self.rc_name.setText(state.get("rc_name", ""))
        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))
        if state.get("sig_time"):
            self.sig_time.setTime(QTime.fromString(state["sig_time"], "HH:mm"))
