# ================================================================
#  A6 FORM PAGE — Section 3 Application by AMHP
#  Mental Health Act 1983 - Form A6 Regulation 4(1)(c)(ii)
#  Single scrollable form layout
# ================================================================

from __future__ import annotations

from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QRadioButton, QButtonGroup, QCheckBox, QPushButton,
    QSizePolicy, QFileDialog, QMessageBox, QGroupBox,
    QToolButton
)


# ================================================================
# TOOLBAR
# ================================================================

class A6Toolbar(QWidget):
    """Toolbar for the A6 Form Page."""

    export_docx = Signal()
    import_file = Signal()
    clear_form = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(60)
        self.setStyleSheet("""
            A6Toolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 8, 16, 8)
        layout.setSpacing(12)

        # Export DOCX button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #1d4ed8; }
            QToolButton:pressed { background: #1e40af; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Import File button
        import_btn = QToolButton()
        import_btn.setText("Import File")
        import_btn.setFixedSize(100, 38)
        import_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #059669; }
            QToolButton:pressed { background: #047857; }
        """)
        import_btn.clicked.connect(self.import_file.emit)
        layout.addWidget(import_btn)

        # Clear Form button
        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #dc2626; }
            QToolButton:pressed { background: #b91c1c; }
        """)
        clear_btn.clicked.connect(self.clear_form.emit)
        layout.addWidget(clear_btn)

        layout.addStretch()


# ================================================================
# MAIN A6 FORM PAGE
# ================================================================

class A6FormPage(QWidget):
    """Page for completing MHA Form A6 - Section 3 AMHP Application for Treatment."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()

        self._setup_ui()
        self._prefill_amhp_details()

    def _load_my_details(self) -> dict:
        """Load clinician details from database."""
        if not self.db:
            return {}

        details = self.db.get_clinician_details()
        if not details:
            return {}

        return {
            "full_name": details[1] or "",
            "email": details[7] or "",
        }

    def _prefill_amhp_details(self):
        """Pre-fill AMHP details from My Details."""
        if self._my_details.get("full_name"):
            self.amhp_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.amhp_email.setText(self._my_details["email"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("""
            QFrame {
                background: #2563eb;
                border-bottom: 1px solid #1d4ed8;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        # Back button
        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: rgba(255,255,255,0.25);
            }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        # Title
        title = QLabel("Form A6 — Section 3 Application by AMHP for Treatment")
        title.setStyleSheet("""
            font-size: 18px;
            font-weight: 700;
            color: white;
        """)
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = A6Toolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        self.toolbar.import_file.connect(self._import_file)
        self.toolbar.clear_form.connect(self._clear_form)
        main_layout.addWidget(self.toolbar)

        # Scrollable form area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        # Build all sections
        self._create_section_1_hospital()
        self._create_section_2_amhp()
        self._create_section_3_patient()
        self._create_section_4_local_authority()
        self._create_section_5_nearest_relative()
        self._create_section_6_interview()
        self._create_section_7_medical_recs()
        self._create_section_8_signature()

        self.form_layout.addStretch()

        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, number: int, title: str) -> QFrame:
        """Create a styled section frame with number and title."""
        frame = QFrame()
        frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)

        # Section header
        header_layout = QHBoxLayout()
        header_layout.setSpacing(12)

        number_badge = QLabel(str(number))
        number_badge.setFixedSize(32, 32)
        number_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        number_badge.setStyleSheet("""
            background: #2563eb;
            color: white;
            font-size: 14px;
            font-weight: 700;
            border-radius: 16px;
        """)
        header_layout.addWidget(number_badge)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 16px;
            font-weight: 600;
            color: #1f2937;
        """)
        header_layout.addWidget(title_lbl)
        header_layout.addStretch()

        layout.addLayout(header_layout)

        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        """Create a styled line edit."""
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus {
                border-color: #2563eb;
            }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 80) -> QTextEdit:
        """Create a styled text edit."""
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMaximumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus {
                border-color: #2563eb;
            }
        """)
        return edit

    def _create_form_row(self, label_text: str, widget: QWidget, label_width: int = 120) -> QHBoxLayout:
        """Create a form row with label and widget."""
        row = QHBoxLayout()
        row.setSpacing(12)
        lbl = QLabel(label_text)
        lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        lbl.setFixedWidth(label_width)
        row.addWidget(lbl)
        row.addWidget(widget, 1)
        return row

    # ----------------------------------------------------------------
    # SECTION 1: Hospital
    # ----------------------------------------------------------------
    def _create_section_1_hospital(self):
        frame = self._create_section_frame(1, "Hospital")
        layout = frame.layout()

        # Hospital name and address on one line
        row = QHBoxLayout()
        row.setSpacing(12)
        self.hospital_name = self._create_line_edit("Hospital name")
        self.hospital_address = self._create_line_edit("Hospital address")
        row.addWidget(self.hospital_name, 1)
        row.addWidget(self.hospital_address, 2)
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 2: AMHP Details
    # ----------------------------------------------------------------
    def _create_section_2_amhp(self):
        frame = self._create_section_frame(2, "AMHP Details")
        layout = frame.layout()

        # All AMHP fields on one line
        row = QHBoxLayout()
        row.setSpacing(12)
        self.amhp_name = self._create_line_edit("Full name")
        self.amhp_address = self._create_line_edit("Address")
        self.amhp_email = self._create_line_edit("Email")
        row.addWidget(self.amhp_name, 1)
        row.addWidget(self.amhp_address, 2)
        row.addWidget(self.amhp_email, 1)
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 3: Patient Details
    # ----------------------------------------------------------------
    def _create_section_3_patient(self):
        frame = self._create_section_frame(3, "Patient Details")
        layout = frame.layout()

        # Patient name and address on one line
        row = QHBoxLayout()
        row.setSpacing(12)
        self.patient_name = self._create_line_edit("Patient full name")
        self.patient_address = self._create_line_edit("Patient address")
        row.addWidget(self.patient_name, 1)
        row.addWidget(self.patient_address, 2)
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 4: Local Authority
    # ----------------------------------------------------------------
    def _create_section_4_local_authority(self):
        frame = self._create_section_frame(4, "Local Authority")
        layout = frame.layout()

        # Row 1: Acting on behalf of + Approved by radios
        row1 = QHBoxLayout()
        row1.setSpacing(12)
        self.local_authority = self._create_line_edit("Acting on behalf of (local social services authority)")
        row1.addWidget(self.local_authority, 2)

        self.approved_by_same = QRadioButton("Approved by same")
        self.approved_by_different = QRadioButton("Different:")
        self.approved_by_same.setChecked(True)
        self.approved_btn_group = QButtonGroup()
        self.approved_btn_group.addButton(self.approved_by_same)
        self.approved_btn_group.addButton(self.approved_by_different)
        row1.addWidget(self.approved_by_same)
        row1.addWidget(self.approved_by_different)

        self.approved_by_authority = self._create_line_edit("Approving authority if different")
        self.approved_by_authority.setEnabled(False)
        row1.addWidget(self.approved_by_authority, 1)
        self.approved_by_different.toggled.connect(self.approved_by_authority.setEnabled)

        layout.addLayout(row1)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 5: Nearest Relative
    # ----------------------------------------------------------------
    def _create_section_5_nearest_relative(self):
        frame = self._create_section_frame(5, "Nearest Relative Consultation")
        layout = frame.layout()

        # Main choice row: Consulted or Not Consulted
        main_row = QHBoxLayout()
        main_row.setSpacing(16)
        self.nr_consulted = QRadioButton("Consulted")
        self.nr_not_consulted = QRadioButton("NOT Consulted")
        self.nr_consulted.setChecked(True)
        self.main_nr_group = QButtonGroup()
        self.main_nr_group.addButton(self.nr_consulted)
        self.main_nr_group.addButton(self.nr_not_consulted)
        main_row.addWidget(self.nr_consulted)
        main_row.addWidget(self.nr_not_consulted)
        main_row.addStretch()
        layout.addLayout(main_row)

        # === CONSULTED SECTION ===
        self.consulted_widget = QWidget()
        consulted_layout = QVBoxLayout(self.consulted_widget)
        consulted_layout.setContentsMargins(0, 8, 0, 0)
        consulted_layout.setSpacing(8)

        # Options row
        opt_row = QHBoxLayout()
        opt_row.setSpacing(16)
        self.nr_option_a = QRadioButton("(a) Is NR")
        self.nr_option_b = QRadioButton("(b) Authorised by court/NR")
        self.nr_option_a.setChecked(True)
        self.option_btn_group = QButtonGroup()
        self.option_btn_group.addButton(self.nr_option_a)
        self.option_btn_group.addButton(self.nr_option_b)
        opt_row.addWidget(self.nr_option_a)
        opt_row.addWidget(self.nr_option_b)
        opt_row.addStretch()
        consulted_layout.addLayout(opt_row)

        # NR name and address on one line
        nr_row = QHBoxLayout()
        nr_row.setSpacing(12)
        self.nr_name = self._create_line_edit("NR full name")
        self.nr_address = self._create_line_edit("NR address")
        nr_row.addWidget(self.nr_name, 1)
        nr_row.addWidget(self.nr_address, 2)
        consulted_layout.addLayout(nr_row)

        layout.addWidget(self.consulted_widget)

        # === NOT CONSULTED SECTION ===
        self.not_consulted_widget = QWidget()
        not_consulted_layout = QVBoxLayout(self.not_consulted_widget)
        not_consulted_layout.setContentsMargins(0, 8, 0, 0)
        not_consulted_layout.setSpacing(8)

        # Options row
        nc_opt_row = QHBoxLayout()
        nc_opt_row.setSpacing(12)
        self.nc_option_a = QRadioButton("(a) Unable to ascertain NR")
        self.nc_option_b = QRadioButton("(b) No NR")
        self.nc_option_c = QRadioButton("(c) Not practicable")
        self.nc_option_a.setChecked(True)
        self.nc_option_group = QButtonGroup()
        self.nc_option_group.addButton(self.nc_option_a)
        self.nc_option_group.addButton(self.nc_option_b)
        self.nc_option_group.addButton(self.nc_option_c)
        nc_opt_row.addWidget(self.nc_option_a)
        nc_opt_row.addWidget(self.nc_option_b)
        nc_opt_row.addWidget(self.nc_option_c)
        nc_opt_row.addStretch()
        not_consulted_layout.addLayout(nc_opt_row)

        # Option C details (only shown when C selected)
        self.nc_option_c_widget = QWidget()
        nc_c_layout = QVBoxLayout(self.nc_option_c_widget)
        nc_c_layout.setContentsMargins(0, 4, 0, 0)
        nc_c_layout.setSpacing(8)

        # NR details row
        nc_nr_row = QHBoxLayout()
        nc_nr_row.setSpacing(12)
        self.nc_nr_name = self._create_line_edit("NR full name")
        self.nc_nr_address = self._create_line_edit("NR address")
        self.nc_c_is_nr = QRadioButton("Is NR")
        self.nc_c_is_authorised = QRadioButton("Authorised")
        self.nc_c_is_nr.setChecked(True)
        self.nc_c_sub_group = QButtonGroup()
        self.nc_c_sub_group.addButton(self.nc_c_is_nr)
        self.nc_c_sub_group.addButton(self.nc_c_is_authorised)
        nc_nr_row.addWidget(self.nc_nr_name, 1)
        nc_nr_row.addWidget(self.nc_nr_address, 2)
        nc_nr_row.addWidget(self.nc_c_is_nr)
        nc_nr_row.addWidget(self.nc_c_is_authorised)
        nc_c_layout.addLayout(nc_nr_row)

        # Reason
        self.nc_reason = self._create_text_edit("Reason why not practicable / unreasonable delay", 80)
        nc_c_layout.addWidget(self.nc_reason)

        not_consulted_layout.addWidget(self.nc_option_c_widget)
        self.nc_option_c_widget.hide()

        layout.addWidget(self.not_consulted_widget)
        self.not_consulted_widget.hide()

        # Connect visibility toggles
        self.nr_consulted.toggled.connect(self._toggle_nr_sections)
        self.nc_option_c.toggled.connect(self.nc_option_c_widget.setVisible)

        self.form_layout.addWidget(frame)

    def _toggle_nr_sections(self, consulted: bool):
        self.consulted_widget.setVisible(consulted)
        self.not_consulted_widget.setVisible(not consulted)

    # ----------------------------------------------------------------
    # SECTION 6: Patient Interview
    # ----------------------------------------------------------------
    def _create_section_6_interview(self):
        frame = self._create_section_frame(6, "Patient Interview")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(12)

        self.last_seen_date = QDateEdit()
        self.last_seen_date.setCalendarPopup(True)
        self.last_seen_date.setDate(QDate.currentDate())
        self.last_seen_date.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        row.addWidget(QLabel("Date seen:"))
        row.addWidget(self.last_seen_date)

        info = QLabel("(within 14 days of application)")
        info.setStyleSheet("font-size: 12px; color: #6b7280;")
        row.addWidget(info)
        row.addStretch()

        layout.addLayout(row)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 7: Medical Recommendations
    # ----------------------------------------------------------------
    def _create_section_7_medical_recs(self):
        frame = self._create_section_frame(7, "Medical Recommendations")
        layout = frame.layout()

        self.no_acquaintance_reason = self._create_text_edit("If neither practitioner had previous acquaintance, explain why (leave blank if N/A)", 80)
        layout.addWidget(self.no_acquaintance_reason)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 8: Signature
    # ----------------------------------------------------------------
    def _create_section_8_signature(self):
        frame = self._create_section_frame(8, "Signature")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(12)

        self.signature_date = QDateEdit()
        self.signature_date.setCalendarPopup(True)
        self.signature_date.setDate(QDate.currentDate())
        self.signature_date.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        row.addWidget(QLabel("Date:"))
        row.addWidget(self.signature_date)

        info = QLabel("(signed manually after printing)")
        info.setStyleSheet("font-size: 12px; color: #6b7280;")
        row.addWidget(info)
        row.addStretch()

        layout.addLayout(row)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _go_back(self):
        self.go_back.emit()

    def _clear_form(self):
        reply = QMessageBox.question(
            self,
            "Clear Form",
            "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            # Clear all fields
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.local_authority.clear()
            self.approved_by_same.setChecked(True)
            self.approved_by_authority.clear()
            self.nr_consulted.setChecked(True)
            self.nr_option_a.setChecked(True)
            self.nr_name.clear()
            self.nr_address.clear()
            self.nc_option_a.setChecked(True)
            self.nc_nr_name.clear()
            self.nc_nr_address.clear()
            self.nc_c_is_nr.setChecked(True)
            self.nc_reason.clear()
            self.last_seen_date.setDate(QDate.currentDate())
            self.no_acquaintance_reason.clear()
            self.signature_date.setDate(QDate.currentDate())

    def _import_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Import File",
            "",
            "All Files (*);;Word Documents (*.docx);;Text Files (*.txt)"
        )
        if file_path:
            QMessageBox.information(self, "Import", f"File selected: {file_path}\n\nData extraction coming soon.")

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Form A6",
            f"Form_A6_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Path to the template
            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_A6_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form A6 template not found.")
                return

            # Open the template
            doc = Document(template_path)

            # Helper to set text in a paragraph
            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    run = para.add_run(new_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            # Helper to add strikethrough to a paragraph
            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            paragraphs = doc.paragraphs

            # Data goes into WHITESPACE paragraphs below instruction paragraphs
            # Para 3: Hospital data (below para 2 instruction)
            # Para 5: AMHP data (below para 4 instruction)
            # Para 7: Patient data (below para 6 instruction)
            # Para 10: Local authority (below para 9 instruction)
            # Para 14: Approved authority (below para 13 instruction)
            # Para 18: NR option a (below para 17 instruction)
            # Para 21: NR option b (below para 20 instruction)
            # Para 28: NR not consulted c (below para 27 instruction)
            # Para 40: Date seen (below para 39 instruction)
            # Para 45: No acquaintance reason (below para 44 instruction)

            # Get data from form fields
            hospital_text = self.hospital_name.text()
            if self.hospital_address.text():
                hospital_text += ", " + self.hospital_address.text()

            amhp_text = self.amhp_name.text()
            if self.amhp_address.text():
                amhp_text += ", " + self.amhp_address.text()
            if self.amhp_email.text():
                amhp_text += ", Email: " + self.amhp_email.text()

            patient_text = self.patient_name.text()
            if self.patient_address.text():
                patient_text += ", " + self.patient_address.text()

            # Fill Hospital - goes in para 3 (whitespace below instruction)
            if hospital_text.strip():
                set_para_text(paragraphs[3], hospital_text)
                highlight_yellow(paragraphs[3])

            # Fill AMHP - goes in para 5 (whitespace below instruction)
            if amhp_text.strip():
                set_para_text(paragraphs[5], amhp_text)
                highlight_yellow(paragraphs[5])

            # Fill Patient - goes in para 7 (whitespace below instruction)
            if patient_text.strip():
                set_para_text(paragraphs[7], patient_text)
                highlight_yellow(paragraphs[7])

            # Fill Local authority - goes in para 10 (whitespace below instruction)
            if self.local_authority.text():
                set_para_text(paragraphs[10], self.local_authority.text())
                highlight_yellow(paragraphs[10])

            # Handle approved by
            if self.approved_by_same.isChecked():
                strikethrough_para(paragraphs[13])
                strikethrough_para(paragraphs[14])
            else:
                strikethrough_para(paragraphs[12])
                if self.approved_by_authority.text():
                    set_para_text(paragraphs[14], self.approved_by_authority.text())
                    highlight_yellow(paragraphs[14])

            # Handle NR consultation
            if self.nr_consulted.isChecked():
                # Consulted - use para 18 or 21, strikethrough 24-36
                nr_text = self.nr_name.text()
                if self.nr_address.text():
                    nr_text += ", " + self.nr_address.text()

                if self.nr_option_a.isChecked():
                    # Option (a) - is NR - data in para 18
                    if nr_text.strip():
                        set_para_text(paragraphs[18], nr_text)
                        highlight_yellow(paragraphs[18])
                    strikethrough_para(paragraphs[20])
                    strikethrough_para(paragraphs[21])
                    strikethrough_para(paragraphs[22])
                else:
                    # Option (b) - authorised - data in para 21
                    if nr_text.strip():
                        set_para_text(paragraphs[21], nr_text)
                        highlight_yellow(paragraphs[21])
                    strikethrough_para(paragraphs[17])
                    strikethrough_para(paragraphs[18])
                    strikethrough_para(paragraphs[19])

                # Strikethrough not consulted section
                for i in range(24, 37):
                    if i < len(paragraphs):
                        strikethrough_para(paragraphs[i])

            else:
                # Not consulted - strikethrough 15-23, use 24-36
                for i in range(15, 24):
                    if i < len(paragraphs):
                        strikethrough_para(paragraphs[i])

                if self.nc_option_a.isChecked():
                    # Unable to ascertain
                    strikethrough_para(paragraphs[26])
                    for i in range(27, 37):
                        if i < len(paragraphs):
                            strikethrough_para(paragraphs[i])
                elif self.nc_option_b.isChecked():
                    # No NR
                    strikethrough_para(paragraphs[25])
                    for i in range(27, 37):
                        if i < len(paragraphs):
                            strikethrough_para(paragraphs[i])
                else:
                    # Option C - not practicable - data in para 28
                    strikethrough_para(paragraphs[25])
                    strikethrough_para(paragraphs[26])

                    nc_nr_text = self.nc_nr_name.text()
                    if self.nc_nr_address.text():
                        nc_nr_text += ", " + self.nc_nr_address.text()

                    if nc_nr_text.strip():
                        set_para_text(paragraphs[28], nc_nr_text)
                        highlight_yellow(paragraphs[28])

                    # Sub-option (i) or (ii)
                    if self.nc_c_is_nr.isChecked():
                        strikethrough_para(paragraphs[31])
                    else:
                        strikethrough_para(paragraphs[30])

                    # Reason - goes in para 34 (whitespace below)
                    if self.nc_reason.toPlainText():
                        set_para_text(paragraphs[34], self.nc_reason.toPlainText())
                        highlight_yellow(paragraphs[34])

            # Date last seen - goes in para 40 (whitespace below instruction)
            last_seen = self.last_seen_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[40], last_seen)
            highlight_yellow(paragraphs[40])

            # No acquaintance reason - goes in para 45 (whitespace below)
            if self.no_acquaintance_reason.toPlainText():
                set_para_text(paragraphs[45], self.no_acquaintance_reason.toPlainText())
                highlight_yellow(paragraphs[45])

            # Signature date (para 49)
            sig_date = self.signature_date.date().toString("dd MMMM yyyy")
            for run in paragraphs[49].runs:
                run.text = ""
            paragraphs[49].add_run(f"Signed                                                                Date {sig_date}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form A6 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        """Get current form state for saving."""
        return {
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.text(),
            "amhp_email": self.amhp_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "local_authority": self.local_authority.text(),
            "approved_by_same": self.approved_by_same.isChecked(),
            "approved_by_authority": self.approved_by_authority.text(),
            "nr_consulted": self.nr_consulted.isChecked(),
            "nr_option_a": self.nr_option_a.isChecked(),
            "nr_name": self.nr_name.text(),
            "nr_address": self.nr_address.text(),
            "nc_option_a": self.nc_option_a.isChecked(),
            "nc_option_b": self.nc_option_b.isChecked(),
            "nc_option_c": self.nc_option_c.isChecked(),
            "nc_nr_name": self.nc_nr_name.text(),
            "nc_nr_address": self.nc_nr_address.text(),
            "nc_c_is_nr": self.nc_c_is_nr.isChecked(),
            "nc_reason": self.nc_reason.toPlainText(),
            "last_seen_date": self.last_seen_date.date().toString("yyyy-MM-dd"),
            "no_acquaintance_reason": self.no_acquaintance_reason.toPlainText(),
            "signature_date": self.signature_date.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        """Load form state."""
        if not state:
            return
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setText(state.get("amhp_address", ""))
        self.amhp_email.setText(state.get("amhp_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.local_authority.setText(state.get("local_authority", ""))
        if state.get("approved_by_same", True):
            self.approved_by_same.setChecked(True)
        else:
            self.approved_by_different.setChecked(True)
        self.approved_by_authority.setText(state.get("approved_by_authority", ""))
        if state.get("nr_consulted", True):
            self.nr_consulted.setChecked(True)
        else:
            self.nr_not_consulted.setChecked(True)
        if state.get("nr_option_a", True):
            self.nr_option_a.setChecked(True)
        else:
            self.nr_option_b.setChecked(True)
        self.nr_name.setText(state.get("nr_name", ""))
        self.nr_address.setText(state.get("nr_address", ""))
        if state.get("nc_option_a", True):
            self.nc_option_a.setChecked(True)
        elif state.get("nc_option_b", False):
            self.nc_option_b.setChecked(True)
        elif state.get("nc_option_c", False):
            self.nc_option_c.setChecked(True)
        self.nc_nr_name.setText(state.get("nc_nr_name", ""))
        self.nc_nr_address.setText(state.get("nc_nr_address", ""))
        if state.get("nc_c_is_nr", True):
            self.nc_c_is_nr.setChecked(True)
        else:
            self.nc_c_is_authorised.setChecked(True)
        self.nc_reason.setPlainText(state.get("nc_reason", ""))
        if state.get("last_seen_date"):
            self.last_seen_date.setDate(QDate.fromString(state["last_seen_date"], "yyyy-MM-dd"))
        self.no_acquaintance_reason.setPlainText(state.get("no_acquaintance_reason", ""))
        if state.get("signature_date"):
            self.signature_date.setDate(QDate.fromString(state["signature_date"], "yyyy-MM-dd"))
