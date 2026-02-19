# ================================================================
#  CTO4 FORM PAGE — Record of Patient's Detention After Recall
#  Mental Health Act 1983 - Form CTO4 Regulation 6(3)(d)
#  Section 17E — Community treatment order: record of detention
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QDateEdit, QTimeEdit,
    QPushButton, QFileDialog, QMessageBox, QToolButton
)


class CTO4FormPage(QWidget):
    """Page for completing MHA Form CTO4 - Record of Detention After Recall."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.signatory_name.setText(self._my_details["full_name"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #7c3aed; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO4 — Record of Detention After Recall")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        toolbar = QWidget()
        toolbar.setFixedHeight(60)
        toolbar.setStyleSheet("background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12);")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 8, 16, 8)
        tb_layout.setSpacing(12)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #7c3aed;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #6d28d9; }
        """)
        export_btn.clicked.connect(self._export_docx)
        tb_layout.addWidget(export_btn)

        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #dc2626; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        tb_layout.addWidget(clear_btn)
        tb_layout.addStretch()

        main_layout.addWidget(toolbar)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        self._build_form()

        self.form_layout.addStretch()
        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, title: str) -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 12px; }")
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("font-size: 16px; font-weight: 700; color: #7c3aed;")
        layout.addWidget(title_lbl)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus { border-color: #7c3aed; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return date_edit

    def _create_time_edit(self) -> QTimeEdit:
        time_edit = QTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("QTimeEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return time_edit

    def _build_form(self):
        # Patient Details
        frame1 = self._create_section_frame("Patient Details")
        layout1 = frame1.layout()

        row1 = QHBoxLayout()
        row1.setSpacing(12)
        self.patient_name = self._create_line_edit("Patient full name")
        row1.addWidget(self.patient_name, 1)
        self.patient_address = self._create_line_edit("Patient address")
        row1.addWidget(self.patient_address, 2)
        layout1.addLayout(row1)

        self.form_layout.addWidget(frame1)

        # Hospital Details
        frame_hosp = self._create_section_frame("Hospital Details")
        layout_hosp = frame_hosp.layout()

        hosp_row = QHBoxLayout()
        hosp_row.setSpacing(12)
        self.hospital_name = self._create_line_edit("Hospital name")
        hosp_row.addWidget(self.hospital_name, 1)
        self.hospital_address = self._create_line_edit("Hospital address")
        hosp_row.addWidget(self.hospital_address, 2)
        layout_hosp.addLayout(hosp_row)

        self.form_layout.addWidget(frame_hosp)

        # Detention Details
        frame2 = self._create_section_frame("Detention Details")
        layout2 = frame2.layout()

        info = QLabel("The patient arrived at the hospital in pursuance of a notice recalling them under section 17E:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 13px; color: #374151;")
        layout2.addWidget(info)

        row2 = QHBoxLayout()
        row2.setSpacing(20)

        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row2.addWidget(date_lbl)
        self.detention_date = self._create_date_edit()
        self.detention_date.setFixedWidth(140)
        row2.addWidget(self.detention_date)

        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row2.addWidget(time_lbl)
        self.detention_time = self._create_time_edit()
        self.detention_time.setFixedWidth(100)
        row2.addWidget(self.detention_time)

        row2.addStretch()
        layout2.addLayout(row2)

        self.form_layout.addWidget(frame2)

        # Signatory Details
        frame3 = self._create_section_frame("Signatory Details")
        layout3 = frame3.layout()

        self.signatory_name = self._create_line_edit("Name (on behalf of hospital managers)")
        layout3.addWidget(self.signatory_name)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(20)

        sig_date_lbl = QLabel("Date:")
        sig_date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_date_lbl)
        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        sig_row.addWidget(self.sig_date)

        sig_time_lbl = QLabel("Time:")
        sig_time_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_time_lbl)
        self.sig_time = self._create_time_edit()
        self.sig_time.setFixedWidth(100)
        sig_row.addWidget(self.sig_time)

        sig_row.addStretch()
        layout3.addLayout(sig_row)

        self.form_layout.addWidget(frame3)

    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.patient_name.clear()
            self.patient_address.clear()
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.detention_date.setDate(QDate.currentDate())
            self.detention_time.setTime(QTime.currentTime())
            self.signatory_name.clear()
            self.sig_date.setDate(QDate.currentDate())
            self.sig_time.setTime(QTime.currentTime())

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form CTO4",
            f"Form_CTO4_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_CTO4_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO4 template not found.")
                return

            doc = Document(template_path)

            def add_yellow_highlight(run):
                """Add yellow highlight to a run."""
                rPr = run._element.get_or_add_rPr()
                shd = rPr.find(qn('w:shd'))
                if shd is None:
                    shd = OxmlElement('w:shd')
                    rPr.append(shd)
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFFCC')

            def replace_in_paragraph(para, old_text, new_text, highlight=True):
                """Replace text in paragraph, preserving structure."""
                full_text = para.text
                if old_text not in full_text:
                    return False
                # Find which run contains the text
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        if highlight:
                            add_yellow_highlight(run)
                        return True
                # If spread across runs, rebuild
                if old_text in full_text:
                    new_full = full_text.replace(old_text, new_text)
                    for i, run in enumerate(para.runs):
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = new_full
                        if highlight:
                            add_yellow_highlight(para.runs[0])
                    return True
                return False

            # Build data strings
            patient_text = self.patient_name.text().strip()
            if self.patient_address.text().strip():
                patient_text += ", " + self.patient_address.text().strip()

            hospital_text = self.hospital_name.text().strip()
            if self.hospital_address.text().strip():
                hospital_text += ", " + self.hospital_address.text().strip()

            det_date = self.detention_date.date().toString("dd/MM/yyyy")
            det_time = self.detention_time.time().toString("HH:mm")
            sig_name = self.signatory_name.text().strip()

            # Process paragraphs - simple bracket replacement with yellow
            for i, para in enumerate(doc.paragraphs):
                text = para.text

                # Patient name and address - keep label, fill blank line below (para 3)
                # Para 2 is the label "[PRINT full name and address of patient]" - keep it
                # Para 3 is blank - fill with patient data
                if i == 3 and text.strip() == "":
                    val = f"[{patient_text}]" if patient_text else "[   ]"
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = val
                        add_yellow_highlight(para.runs[0])
                    else:
                        run = para.add_run(val)
                        add_yellow_highlight(run)

                # Hospital name and address - keep label, fill blank line below (para 6)
                # Para 5 has "[full name and address of hospital]" - keep it as label
                # Para 6 is blank - fill with hospital data
                if i == 6 and text.strip() == "":
                    val = f"[{hospital_text}]" if hospital_text else "[   ]"
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = val
                        add_yellow_highlight(para.runs[0])
                    else:
                        run = para.add_run(val)
                        add_yellow_highlight(run)

                # Para 7 has "[enter date and time...]" - keep it as label, don't modify
                # Date and time values go in para 8 which has "Date ... Time ..." row

                # Row with Date and Time labels (para 8) - add brackets with values
                if text.strip().startswith("Date") and "Time" in text and len(text.strip()) < 100:
                    # Clear and rebuild: Date [value]    Time [value]
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = "Date "
                    date_run = para.add_run(f"[{det_date}]" if det_date else "[   ]")
                    add_yellow_highlight(date_run)
                    para.add_run("                                                       Time ")
                    time_run = para.add_run(f"[{det_time}]" if det_time else "[   ]")
                    add_yellow_highlight(time_run)

                # PRINT NAME row - add bracket with value
                if text.strip().startswith("PRINT NAME"):
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = "PRINT NAME "
                    name_run = para.add_run(f"[{sig_name}]" if sig_name else "[   ]")
                    add_yellow_highlight(name_run)

                # Date row (signatory) - single Date label
                if text.strip() == "Date" or (text.strip().startswith("Date") and "Time" not in text and len(text.strip()) < 80):
                    # Skip if already processed above
                    if "Time" not in text:
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = "Date "
                        date_run = para.add_run(f"[{det_date}]" if det_date else "[   ]")
                        add_yellow_highlight(date_run)

                # Time row (signatory)
                if text.strip() == "Time" or (text.strip().startswith("Time") and "Date" not in text and len(text.strip()) < 80):
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = "Time "
                    time_run = para.add_run(f"[{det_time}]" if det_time else "[   ]")
                    add_yellow_highlight(time_run)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO4 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    def get_state(self) -> dict:
        return {
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "detention_date": self.detention_date.date().toString("yyyy-MM-dd"),
            "detention_time": self.detention_time.time().toString("HH:mm"),
            "signatory_name": self.signatory_name.text(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
            "sig_time": self.sig_time.time().toString("HH:mm"),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        if state.get("detention_date"):
            self.detention_date.setDate(QDate.fromString(state["detention_date"], "yyyy-MM-dd"))
        if state.get("detention_time"):
            self.detention_time.setTime(QTime.fromString(state["detention_time"], "HH:mm"))
        self.signatory_name.setText(state.get("signatory_name", ""))
        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))
        if state.get("sig_time"):
            self.sig_time.setTime(QTime.fromString(state["sig_time"], "HH:mm"))
