# ================================================================
#  CTO5 FORM PAGE — Revocation of Community Treatment Order
#  Mental Health Act 1983 - Form CTO5 Regulation 6(8)(a) and (b)
#  Section 17F(4) — Revocation of community treatment order
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QToolButton, QRadioButton, QButtonGroup,
    QComboBox, QSpinBox, QCompleter, QStyleFactory, QSlider
)

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


class CTO5FormPage(QWidget):
    """Page for completing MHA Form CTO5 - Revocation of CTO."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean",
        "Asian",
        "Caucasian",
        "Middle Eastern",
        "Mixed Race",
        "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #ea580c; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO5 — Revocation of Community Treatment Order")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        toolbar = QWidget()
        toolbar.setFixedHeight(60)
        toolbar.setStyleSheet("background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12);")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 8, 16, 8)
        tb_layout.setSpacing(12)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #ea580c;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #c2410c; }
        """)
        export_btn.clicked.connect(self._export_docx)
        tb_layout.addWidget(export_btn)

        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #dc2626; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        tb_layout.addWidget(clear_btn)
        tb_layout.addStretch()

        main_layout.addWidget(toolbar)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        self._create_section_1_demographics()
        self._create_section_2_patient_rc()
        self._create_section_3_reasons()
        self._create_section_4_amhp()
        self._create_section_5_revocation()

        self.form_layout.addStretch()
        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, number: int, title: str, color: str = "#ea580c") -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)

        header_layout = QHBoxLayout()
        header_layout.setSpacing(12)

        number_badge = QLabel(str(number))
        number_badge.setFixedSize(32, 32)
        number_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        number_badge.setStyleSheet(f"""
            background: {color};
            color: white;
            font-size: 14px;
            font-weight: 700;
            border-radius: 16px;
        """)
        header_layout.addWidget(number_badge)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1f2937;")
        header_layout.addWidget(title_lbl)
        header_layout.addStretch()

        layout.addLayout(header_layout)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus { border-color: #ea580c; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 100) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMinimumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #ea580c; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return date_edit

    def _create_time_edit(self) -> QTimeEdit:
        time_edit = QTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("QTimeEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return time_edit

    # ----------------------------------------------------------------
    # SECTION 1: Demographics
    # ----------------------------------------------------------------
    def _create_section_1_demographics(self):
        frame = self._create_section_frame(1, "Patient Demographics")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(16)

        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row.addWidget(age_lbl)

        self.age_spin = QSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setFixedWidth(70)
        self.age_spin.setStyleSheet("QSpinBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 13px; }")
        row.addWidget(self.age_spin)

        row.addSpacing(20)

        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("Male")
        self.gender_female = QRadioButton("Female")
        self.gender_other = QRadioButton("Other")
        self.gender_group.addButton(self.gender_male, 0)
        self.gender_group.addButton(self.gender_female, 1)
        self.gender_group.addButton(self.gender_other, 2)

        row.addWidget(self.gender_male)
        row.addWidget(self.gender_female)
        row.addWidget(self.gender_other)

        row.addSpacing(20)

        self.ethnicity_combo = QComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(200)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 13px; }")
        row.addWidget(self.ethnicity_combo)

        row.addStretch()
        layout.addLayout(row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 2: Patient & RC Details
    # ----------------------------------------------------------------
    def _create_section_2_patient_rc(self):
        frame = self._create_section_frame(2, "PART 1 — Responsible Clinician")
        layout = frame.layout()

        # RC Details
        rc_lbl = QLabel("RC Details:")
        rc_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151;")
        layout.addWidget(rc_lbl)

        row1 = QHBoxLayout()
        row1.setSpacing(12)
        self.rc_name = self._create_line_edit("RC full name")
        row1.addWidget(self.rc_name, 1)
        self.rc_address = self._create_line_edit("Address")
        row1.addWidget(self.rc_address, 2)
        self.rc_email = self._create_line_edit("Email")
        row1.addWidget(self.rc_email, 1)
        layout.addLayout(row1)

        # Patient Details
        pt_lbl = QLabel("Patient Details:")
        pt_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151; margin-top: 8px;")
        layout.addWidget(pt_lbl)

        row2 = QHBoxLayout()
        row2.setSpacing(12)
        self.patient_name = self._create_line_edit("Patient full name")
        row2.addWidget(self.patient_name, 1)
        self.patient_address = self._create_line_edit("Patient address")
        row2.addWidget(self.patient_address, 2)
        layout.addLayout(row2)

        # Hospital
        hosp_lbl = QLabel("Detained in:")
        hosp_lbl.setStyleSheet("font-size: 12px; font-weight: 500; color: #374151; margin-top: 8px;")
        layout.addWidget(hosp_lbl)

        self.hospital = self._create_line_edit("Hospital name and address")
        layout.addWidget(self.hospital)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 3: Reasons for Revocation (with clinical controls)
    # ----------------------------------------------------------------
    def _create_section_3_reasons(self):
        frame = self._create_section_frame(3, "Reasons for Revocation")
        layout = frame.layout()

        split_layout = QHBoxLayout()
        split_layout.setSpacing(16)

        # === LEFT: Reasons Text Area ===
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(6)

        info = QLabel("Click options on the right to auto-generate reasons:")
        info.setStyleSheet("font-size: 10px; color: #6b7280; padding: 4px; background: #fff7ed; border-radius: 4px;")
        left_layout.addWidget(info)

        self.reasons = QTextEdit()
        self.reasons.setPlaceholderText("Reasons for revocation will be generated here...")
        self.reasons.setMinimumHeight(400)
        self.reasons.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px; font-size: 12px; }
            QTextEdit:focus { border-color: #ea580c; }
        """)
        left_layout.addWidget(self.reasons)

        # RC Signature
        sig_row = QHBoxLayout()
        sig_row.setSpacing(12)
        sig_lbl = QLabel("RC Signature Date:")
        sig_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(140)
        sig_row.addWidget(self.rc_sig_date)
        sig_row.addStretch()
        left_layout.addLayout(sig_row)

        split_layout.addWidget(left_container, 3)

        # === RIGHT: Controls Panel ===
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        right_scroll.setFixedWidth(400)

        right_container = QWidget()
        right_container.setFixedWidth(380)
        right_container.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 8, 0)
        right_layout.setSpacing(10)

        # --- Mental Disorder (ICD-10) ---
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(10, 8, 10, 8)
        md_layout.setSpacing(6)

        md_header = QLabel("Mental Disorder")
        md_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        self.dx_boxes = []
        for i in range(2):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText("Primary diagnosis..." if i == 0 else "Secondary (optional)...")

            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(ICD10_DICT.items(), key=lambda x: x[0].lower()):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})

            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(12)
            combo.setStyleSheet("""
                QComboBox { padding: 5px; font-size: 11px; border: 1px solid #d1d5db; border-radius: 4px; background: white; }
                QComboBox QAbstractItemView { min-width: 300px; }
            """)
            combo.currentIndexChanged.connect(self._update_reasons_text)
            md_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        right_layout.addWidget(md_frame)

        # --- Legal Criteria ---
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(10, 8, 10, 8)
        lc_layout.setSpacing(4)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._update_reasons_text)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._update_reasons_text)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._update_reasons_text)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(100)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 11px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 11px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_reasons_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 11px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 11px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._update_reasons_text)
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._update_reasons_text)
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 11px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 11px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._update_reasons_text)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 11px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # Self
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_hist_harm)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._update_reasons_text)
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 11px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence")
        self.others_hist_violence.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_hist_verbal)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 10px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence")
        self.others_curr_violence.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 10px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._update_reasons_text)
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        right_layout.addWidget(lc_frame)

        # --- Revocation Reasons ---
        rr_frame = QFrame()
        rr_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; }")
        rr_layout = QVBoxLayout(rr_frame)
        rr_layout.setContentsMargins(10, 8, 10, 8)
        rr_layout.setSpacing(4)

        rr_header = QLabel("Revocation Because")
        rr_header.setStyleSheet("font-size: 11px; font-weight: 700; color: #991b1b;")
        rr_layout.addWidget(rr_header)

        self.cto_failed_cb = QCheckBox("CTO no longer effective")
        self.cto_failed_cb.setStyleSheet("font-size: 11px; color: #374151;")
        self.cto_failed_cb.toggled.connect(self._update_reasons_text)
        rr_layout.addWidget(self.cto_failed_cb)

        self.relapse_cb = QCheckBox("Relapsed despite CTO")
        self.relapse_cb.setStyleSheet("font-size: 11px; color: #374151;")
        self.relapse_cb.toggled.connect(self._update_reasons_text)
        rr_layout.addWidget(self.relapse_cb)

        self.inpatient_cb = QCheckBox("Requires inpatient treatment")
        self.inpatient_cb.setStyleSheet("font-size: 11px; color: #374151;")
        self.inpatient_cb.toggled.connect(self._update_reasons_text)
        rr_layout.addWidget(self.inpatient_cb)

        right_layout.addWidget(rr_frame)
        right_layout.addStretch()

        right_scroll.setWidget(right_container)
        split_layout.addWidget(right_scroll)

        layout.addLayout(split_layout)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 4: AMHP Agreement
    # ----------------------------------------------------------------
    def _create_section_4_amhp(self):
        frame = self._create_section_frame(4, "PART 2 — Approved Mental Health Professional", "#059669")
        layout = frame.layout()

        row1 = QHBoxLayout()
        row1.setSpacing(12)
        self.amhp_name = self._create_line_edit("AMHP full name")
        row1.addWidget(self.amhp_name, 1)
        self.amhp_address = self._create_line_edit("AMHP address")
        row1.addWidget(self.amhp_address, 2)
        self.amhp_email = self._create_line_edit("Email")
        row1.addWidget(self.amhp_email, 1)
        layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.setSpacing(12)
        self.amhp_authority = self._create_line_edit("Acting on behalf of (Authority)")
        row2.addWidget(self.amhp_authority, 1)
        self.amhp_approved_by = self._create_line_edit("Approved by (if different)")
        row2.addWidget(self.amhp_approved_by, 1)
        layout.addLayout(row2)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(12)
        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.amhp_sig_date = self._create_date_edit()
        self.amhp_sig_date.setFixedWidth(140)
        sig_row.addWidget(self.amhp_sig_date)
        sig_row.addStretch()
        layout.addLayout(sig_row)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 5: Revocation Order
    # ----------------------------------------------------------------
    def _create_section_5_revocation(self):
        frame = self._create_section_frame(5, "PART 3 — Revocation Order", "#dc2626")
        layout = frame.layout()

        info = QLabel("I exercise my power under section 17F(4) to revoke the community treatment order:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 13px; color: #374151;")
        layout.addWidget(info)

        row = QHBoxLayout()
        row.setSpacing(20)

        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row.addWidget(date_lbl)
        self.revoke_date = self._create_date_edit()
        self.revoke_date.setFixedWidth(140)
        row.addWidget(self.revoke_date)

        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        row.addWidget(time_lbl)
        self.revoke_time = self._create_time_edit()
        self.revoke_time.setFixedWidth(100)
        row.addWidget(self.revoke_time)

        row.addStretch()
        layout.addLayout(row)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(12)
        sig_lbl = QLabel("RC Signature Date:")
        sig_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.final_sig_date = self._create_date_edit()
        self.final_sig_date.setFixedWidth(140)
        sig_row.addWidget(self.final_sig_date)
        sig_row.addStretch()
        layout.addLayout(sig_row)

        warning = QLabel("This revocation is NOT VALID unless all three parts are completed and signed.")
        warning.setWordWrap(True)
        warning.setStyleSheet("font-size: 12px; color: #dc2626; font-weight: 600; padding: 8px; background: #fef2f2; border-radius: 6px; margin-top: 8px;")
        layout.addWidget(warning)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # Toggle handlers
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_reasons_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_reasons_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_reasons_text()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_reasons_text()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_reasons_text()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_reasons_text()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_reasons_text()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm,
                       self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
                cb.setChecked(False)
        self._update_reasons_text()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            for cb in [self.others_hist_violence, self.others_hist_verbal,
                       self.others_curr_violence, self.others_curr_verbal]:
                cb.setChecked(False)
        self._update_reasons_text()

    # ----------------------------------------------------------------
    # Text generation
    # ----------------------------------------------------------------
    def _update_reasons_text(self):
        self.reasons.setPlainText(self._generate_reasons_text())

    def _generate_reasons_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        paragraphs = []

        # Para 1: Demographics + Diagnosis
        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", "").replace("Mixed ", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    diagnoses.append(f"{dx} ({icd})" if icd else dx)

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for the patient to receive medical treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for the patient to receive medical treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for the patient to receive medical treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        # Para 2: Necessity
        para2_parts = []
        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("protection of others")

        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"Inpatient treatment is necessary for {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"Inpatient treatment is necessary for {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"Inpatient treatment is necessary for {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")

        # Health details
        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_items = []
            if self.poor_compliance_cb.isChecked():
                mh_items.append("poor compliance with treatment")
            if self.limited_insight_cb.isChecked():
                mh_items.append(f"limited insight into {p['pos_l']} illness")
            if mh_items:
                para2_parts.append(f"Regarding {p['pos_l']} health, I am concerned about {' and '.join(mh_items)}.")

        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health: {details}.")

        # Safety - Self
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            if self.gender_male.isChecked():
                reflexive = "himself"
            elif self.gender_female.isChecked():
                reflexive = "herself"
            else:
                reflexive = "themselves"

            risk_types = [
                ("self neglect", self.self_hist_neglect.isChecked(), self.self_curr_neglect.isChecked()),
                (f"placing of {reflexive} in risky situations", self.self_hist_risky.isChecked(), self.self_curr_risky.isChecked()),
                ("self harm", self.self_hist_harm.isChecked(), self.self_curr_harm.isChecked()),
            ]

            both_items = []
            hist_only = []
            curr_only = []

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            self_text = f"With respect to {p['pos_l']} own safety I am concerned about"
            parts = []
            if both_items:
                if len(both_items) == 1:
                    parts.append(f"historical and current {both_items[0]}")
                else:
                    parts.append(f"historical and current {', '.join(both_items[:-1])}, and of {both_items[-1]}")
            if hist_only:
                if len(hist_only) == 1:
                    parts.append(f"historical {hist_only[0]}")
                else:
                    parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
            if curr_only:
                if len(curr_only) == 1:
                    parts.append(f"current {curr_only[0]}")
                else:
                    parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")

            if parts:
                self_text += " " + ", and ".join(parts) + "."
                para2_parts.append(self_text)

        # Safety - Others
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            risk_types = [
                ("violence to others", self.others_hist_violence.isChecked(), self.others_curr_violence.isChecked()),
                ("verbal aggression", self.others_hist_verbal.isChecked(), self.others_curr_verbal.isChecked()),
            ]

            both_items = []
            hist_only = []
            curr_only = []

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            others_text = f"Regarding risk to others, I am concerned about the risk of"
            parts = []
            if both_items:
                if len(both_items) == 1:
                    parts.append(f"{both_items[0]} which is both historical and current")
                else:
                    parts.append(f"{', '.join(both_items[:-1])}, and {both_items[-1]} which are both historical and current")
            if hist_only:
                if len(hist_only) == 1:
                    parts.append(f"historical {hist_only[0]}")
                else:
                    parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
            if curr_only:
                if len(curr_only) == 1:
                    parts.append(f"current {curr_only[0]}")
                else:
                    parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")

            if parts:
                others_text += " " + " and of ".join(parts) + "."
                para2_parts.append(others_text)

        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        # Para 3: Revocation reasons
        para3_parts = []
        revoke_reasons = []
        if self.cto_failed_cb.isChecked():
            revoke_reasons.append("the CTO is no longer effective in managing the patient's care")
        if self.relapse_cb.isChecked():
            revoke_reasons.append(f"{p['subj_l']} has relapsed despite being on a CTO")
        if self.inpatient_cb.isChecked():
            revoke_reasons.append(f"{p['subj_l']} requires inpatient treatment that cannot be provided under the CTO")

        if revoke_reasons:
            para3_parts.append(f"The CTO should be revoked because {', and '.join(revoke_reasons)}.")

        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        return "\n\n".join(paragraphs)

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.rc_name.clear()
            self.rc_address.clear()
            self.rc_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.hospital.clear()
            for combo in self.dx_boxes:
                combo.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            self.cto_failed_cb.setChecked(False)
            self.relapse_cb.setChecked(False)
            self.inpatient_cb.setChecked(False)
            self.reasons.clear()
            self.rc_sig_date.setDate(QDate.currentDate())
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_email.clear()
            self.amhp_authority.clear()
            self.amhp_approved_by.clear()
            self.amhp_sig_date.setDate(QDate.currentDate())
            self.revoke_date.setDate(QDate.currentDate())
            self.revoke_time.setTime(QTime.currentTime())
            self.final_sig_date.setDate(QDate.currentDate())
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form CTO5",
            f"Form_CTO5_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_CTO5_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO5 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # RC details (para 5)
            rc_text = self.rc_name.text()
            if self.rc_address.text():
                rc_text += ", " + self.rc_address.text()
            if self.rc_email.text():
                rc_text += ", Email: " + self.rc_email.text()
            if rc_text.strip():
                set_para_text(paragraphs[5], rc_text)
                highlight_yellow(paragraphs[5])

            # Patient (para 7)
            patient_text = self.patient_name.text()
            if self.patient_address.text():
                patient_text += ", " + self.patient_address.text()
            if patient_text.strip():
                set_para_text(paragraphs[7], patient_text)
                highlight_yellow(paragraphs[7])

            # Hospital (para 9)
            if self.hospital.text().strip():
                set_para_text(paragraphs[9], self.hospital.text())
                highlight_yellow(paragraphs[9])

            # Necessity (strikethrough unselected)
            if not self.health_cb.isChecked():
                strikethrough_para(paragraphs[15])
            if not (self.safety_cb.isChecked() and self.self_harm_cb.isChecked()):
                strikethrough_para(paragraphs[16])
            if not (self.safety_cb.isChecked() and self.others_cb.isChecked()):
                strikethrough_para(paragraphs[17])

            # Reasons (para 23-25)
            if self.reasons.toPlainText().strip():
                set_para_text(paragraphs[23], self.reasons.toPlainText())
                highlight_yellow(paragraphs[23])

            # Part 1 sig (para 28-29)
            sig_date = self.rc_sig_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[29], sig_date)

            # AMHP details (para 33)
            amhp_text = self.amhp_name.text()
            if self.amhp_address.text():
                amhp_text += ", " + self.amhp_address.text()
            if self.amhp_email.text():
                amhp_text += ", Email: " + self.amhp_email.text()
            if amhp_text.strip():
                set_para_text(paragraphs[33], amhp_text)
                highlight_yellow(paragraphs[33])

            # AMHP authority (para 35)
            if self.amhp_authority.text().strip():
                set_para_text(paragraphs[35], self.amhp_authority.text())
                highlight_yellow(paragraphs[35])

            # AMHP sig (para 45)
            amhp_sig = self.amhp_sig_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[45], amhp_sig)

            # Part 3 - revocation date/time (para 48)
            rev_date = self.revoke_date.date().toString("dd MMMM yyyy")
            rev_time = self.revoke_time.time().toString("HH:mm")
            set_para_text(paragraphs[48], f"Date: {rev_date}                    Time: {rev_time}")

            # Final sig (para 51)
            final_date = self.final_sig_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[51], f"Responsible clinician Date {final_date}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO5 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    def get_state(self) -> dict:
        return {
            "age": self.age_spin.value(),
            "gender": "male" if self.gender_male.isChecked() else "female" if self.gender_female.isChecked() else "other" if self.gender_other.isChecked() else "",
            "ethnicity": self.ethnicity_combo.currentText(),
            "rc_name": self.rc_name.text(),
            "rc_address": self.rc_address.text(),
            "rc_email": self.rc_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "hospital": self.hospital.text(),
            "diagnoses": [combo.currentText() for combo in self.dx_boxes],
            "nature": self.nature_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "health": self.health_cb.isChecked(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "others": self.others_cb.isChecked(),
            "cto_failed": self.cto_failed_cb.isChecked(),
            "relapse": self.relapse_cb.isChecked(),
            "inpatient": self.inpatient_cb.isChecked(),
            "reasons": self.reasons.toPlainText(),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.text(),
            "amhp_email": self.amhp_email.text(),
            "amhp_authority": self.amhp_authority.text(),
            "amhp_approved_by": self.amhp_approved_by.text(),
            "amhp_sig_date": self.amhp_sig_date.date().toString("yyyy-MM-dd"),
            "revoke_date": self.revoke_date.date().toString("yyyy-MM-dd"),
            "revoke_time": self.revoke_time.time().toString("HH:mm"),
            "final_sig_date": self.final_sig_date.date().toString("yyyy-MM-dd"),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.age_spin.setValue(state.get("age", 0))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        ethnicity = state.get("ethnicity", "Ethnicity")
        idx = self.ethnicity_combo.findText(ethnicity)
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_address.setText(state.get("rc_address", ""))
        self.rc_email.setText(state.get("rc_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.hospital.setText(state.get("hospital", ""))
        self.nature_cb.setChecked(state.get("nature", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.health_cb.setChecked(state.get("health", False))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.cto_failed_cb.setChecked(state.get("cto_failed", False))
        self.relapse_cb.setChecked(state.get("relapse", False))
        self.inpatient_cb.setChecked(state.get("inpatient", False))
        self.reasons.setPlainText(state.get("reasons", ""))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setText(state.get("amhp_address", ""))
        self.amhp_email.setText(state.get("amhp_email", ""))
        self.amhp_authority.setText(state.get("amhp_authority", ""))
        self.amhp_approved_by.setText(state.get("amhp_approved_by", ""))
        if state.get("amhp_sig_date"):
            self.amhp_sig_date.setDate(QDate.fromString(state["amhp_sig_date"], "yyyy-MM-dd"))
        if state.get("revoke_date"):
            self.revoke_date.setDate(QDate.fromString(state["revoke_date"], "yyyy-MM-dd"))
        if state.get("revoke_time"):
            self.revoke_time.setTime(QTime.fromString(state["revoke_time"], "HH:mm"))
        if state.get("final_sig_date"):
            self.final_sig_date.setDate(QDate.fromString(state["final_sig_date"], "yyyy-MM-dd"))
