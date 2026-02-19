# ================================================================
#  A2 FORM PAGE — Section 2 Application by AMHP
#  Mental Health Act 1983 - Form A2 Regulation 4(1)(a)(ii)
#  Single scrollable form layout
# ================================================================

from __future__ import annotations

from datetime import datetime
from typing import Optional
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QRadioButton, QButtonGroup, QCheckBox, QPushButton,
    QSizePolicy, QFileDialog, QMessageBox, QGroupBox,
    QFormLayout, QSpacerItem, QToolButton
)


# ================================================================
# TOOLBAR
# ================================================================

class A2Toolbar(QWidget):
    """Toolbar for the A2 Form Page."""

    export_docx = Signal()
    import_file = Signal()
    clear_form = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(60)
        self.setStyleSheet("""
            A2Toolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 8, 16, 8)
        layout.setSpacing(12)

        # Export DOCX button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #1d4ed8; }
            QToolButton:pressed { background: #1e40af; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Import File button
        import_btn = QToolButton()
        import_btn.setText("Import File")
        import_btn.setFixedSize(100, 38)
        import_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #059669; }
            QToolButton:pressed { background: #047857; }
        """)
        import_btn.clicked.connect(self.import_file.emit)
        layout.addWidget(import_btn)

        # Clear Form button
        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #dc2626; }
            QToolButton:pressed { background: #b91c1c; }
        """)
        clear_btn.clicked.connect(self.clear_form.emit)
        layout.addWidget(clear_btn)

        layout.addStretch()


# ================================================================
# MAIN A2 FORM PAGE
# ================================================================

class A2FormPage(QWidget):
    """Page for completing MHA Form A2 - Section 2 AMHP Application."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()

        self._setup_ui()
        self._prefill_amhp_details()

    def _load_my_details(self) -> dict:
        """Load clinician details from database."""
        if not self.db:
            return {}

        details = self.db.get_clinician_details()
        if not details:
            return {}

        return {
            "full_name": details[1] or "",
            "role_title": details[2] or "",
            "discipline": details[3] or "",
            "registration_body": details[4] or "",
            "registration_number": details[5] or "",
            "phone": details[6] or "",
            "email": details[7] or "",
            "team_service": details[8] or "",
            "hospital_org": details[9] or "",
            "ward_department": details[10] or "",
        }

    def _prefill_amhp_details(self):
        """Pre-fill AMHP details from My Details."""
        if self._my_details.get("full_name"):
            self.amhp_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.amhp_email.setText(self._my_details["email"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("""
            QFrame {
                background: #2563eb;
                border-bottom: 1px solid #1d4ed8;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        # Back button
        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: rgba(255,255,255,0.25);
            }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        # Title
        title = QLabel("Form A2 — Section 2 Application by AMHP")
        title.setStyleSheet("""
            font-size: 18px;
            font-weight: 700;
            color: white;
        """)
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = A2Toolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        self.toolbar.import_file.connect(self._import_file)
        self.toolbar.clear_form.connect(self._clear_form)
        main_layout.addWidget(self.toolbar)

        # Scrollable form area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        # Build all sections
        self._create_section_1_hospital()
        self._create_section_2_amhp()
        self._create_section_3_patient()
        self._create_section_4_local_authority()
        self._create_section_5_nearest_relative()
        self._create_section_6_interview()
        self._create_section_7_medical_recs()
        self._create_section_8_signature()

        self.form_layout.addStretch()

        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, number: int, title: str) -> QFrame:
        """Create a styled section frame with number and title."""
        frame = QFrame()
        frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)

        # Section header
        header_layout = QHBoxLayout()
        header_layout.setSpacing(12)

        number_badge = QLabel(str(number))
        number_badge.setFixedSize(32, 32)
        number_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        number_badge.setStyleSheet("""
            background: #2563eb;
            color: white;
            font-size: 14px;
            font-weight: 700;
            border-radius: 16px;
        """)
        header_layout.addWidget(number_badge)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 16px;
            font-weight: 600;
            color: #1f2937;
        """)
        header_layout.addWidget(title_lbl)
        header_layout.addStretch()

        layout.addLayout(header_layout)

        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        """Create a styled line edit."""
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus {
                border-color: #2563eb;
            }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 80) -> QTextEdit:
        """Create a styled text edit."""
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMaximumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus {
                border-color: #2563eb;
            }
        """)
        return edit

    def _create_form_row(self, label_text: str, widget: QWidget, label_width: int = 120) -> QHBoxLayout:
        """Create a form row with label and widget."""
        row = QHBoxLayout()
        row.setSpacing(12)
        lbl = QLabel(label_text)
        lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        lbl.setFixedWidth(label_width)
        row.addWidget(lbl)
        row.addWidget(widget, 1)
        return row

    # ----------------------------------------------------------------
    # SECTION 1: Hospital
    # ----------------------------------------------------------------
    def _create_section_1_hospital(self):
        frame = self._create_section_frame(1, "Hospital")
        layout = frame.layout()

        self.hospital_name = self._create_line_edit("Enter hospital name")
        layout.addLayout(self._create_form_row("Hospital Name:", self.hospital_name))

        self.hospital_address = self._create_text_edit("Enter hospital address")
        layout.addLayout(self._create_form_row("Address:", self.hospital_address))

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 2: AMHP Details
    # ----------------------------------------------------------------
    def _create_section_2_amhp(self):
        frame = self._create_section_frame(2, "AMHP Details")
        layout = frame.layout()

        self.amhp_name = self._create_line_edit("Enter your full name")
        layout.addLayout(self._create_form_row("Full Name:", self.amhp_name))

        self.amhp_address = self._create_text_edit("Enter your address")
        layout.addLayout(self._create_form_row("Address:", self.amhp_address))

        self.amhp_email = self._create_line_edit("Enter email address")
        layout.addLayout(self._create_form_row("Email:", self.amhp_email))

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 3: Patient Details
    # ----------------------------------------------------------------
    def _create_section_3_patient(self):
        frame = self._create_section_frame(3, "Patient Details")
        layout = frame.layout()

        self.patient_name = self._create_line_edit("Enter patient's full name")
        layout.addLayout(self._create_form_row("Full Name:", self.patient_name))

        self.patient_address = self._create_text_edit("Enter patient's address")
        layout.addLayout(self._create_form_row("Address:", self.patient_address))

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 4: Local Authority
    # ----------------------------------------------------------------
    def _create_section_4_local_authority(self):
        frame = self._create_section_frame(4, "Local Authority")
        layout = frame.layout()

        self.local_authority = self._create_line_edit("Enter local social services authority")
        layout.addLayout(self._create_form_row("Authority:", self.local_authority, 140))

        # Approved by group
        group = QGroupBox("Approved by")
        group.setStyleSheet("""
            QGroupBox {
                font-weight: 600;
                font-size: 13px;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                margin-top: 12px;
                padding-top: 20px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 8px;
            }
        """)
        group_layout = QVBoxLayout(group)

        self.approved_by_same = QRadioButton("That authority (same as above)")
        self.approved_by_different = QRadioButton("Different authority:")
        self.approved_by_same.setChecked(True)

        self.approved_btn_group = QButtonGroup()
        self.approved_btn_group.addButton(self.approved_by_same)
        self.approved_btn_group.addButton(self.approved_by_different)

        group_layout.addWidget(self.approved_by_same)
        group_layout.addWidget(self.approved_by_different)

        self.approved_by_authority = self._create_line_edit("Enter approving authority if different")
        self.approved_by_authority.setEnabled(False)
        group_layout.addWidget(self.approved_by_authority)

        self.approved_by_different.toggled.connect(self.approved_by_authority.setEnabled)

        layout.addWidget(group)
        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 5: Nearest Relative
    # ----------------------------------------------------------------
    def _create_section_5_nearest_relative(self):
        frame = self._create_section_frame(5, "Nearest Relative")
        layout = frame.layout()

        # Known/Unknown
        known_group = QGroupBox("Do you know who the nearest relative is?")
        known_group.setStyleSheet("""
            QGroupBox {
                font-weight: 600;
                font-size: 13px;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                margin-top: 8px;
                padding-top: 20px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 8px;
            }
        """)
        known_layout = QVBoxLayout(known_group)

        self.nr_known = QRadioButton("Yes - I know who the nearest relative is")
        self.nr_unknown = QRadioButton("No - I don't know / patient has no NR")
        self.nr_known.setChecked(True)

        self.known_btn_group = QButtonGroup()
        self.known_btn_group.addButton(self.nr_known)
        self.known_btn_group.addButton(self.nr_unknown)

        known_layout.addWidget(self.nr_known)
        known_layout.addWidget(self.nr_unknown)
        layout.addWidget(known_group)

        # NR Details (when known)
        self.nr_details_widget = QWidget()
        nr_details_layout = QVBoxLayout(self.nr_details_widget)
        nr_details_layout.setContentsMargins(0, 8, 0, 0)

        # Option (a) or (b)
        option_group = QGroupBox("Select option (a) or (b)")
        option_group.setStyleSheet("""
            QGroupBox {
                font-weight: 500;
                font-size: 12px;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding-top: 16px;
            }
        """)
        option_layout = QVBoxLayout(option_group)

        self.nr_option_a = QRadioButton("(a) This person IS the patient's nearest relative")
        self.nr_option_b = QRadioButton("(b) This person has been AUTHORISED to act as NR")
        self.nr_option_a.setChecked(True)

        self.option_btn_group = QButtonGroup()
        self.option_btn_group.addButton(self.nr_option_a)
        self.option_btn_group.addButton(self.nr_option_b)

        option_layout.addWidget(self.nr_option_a)
        option_layout.addWidget(self.nr_option_b)
        nr_details_layout.addWidget(option_group)

        # NR name and address
        self.nr_name = self._create_line_edit("Enter nearest relative's full name")
        nr_details_layout.addLayout(self._create_form_row("NR Name:", self.nr_name))

        self.nr_address = self._create_text_edit("Enter nearest relative's address")
        nr_details_layout.addLayout(self._create_form_row("NR Address:", self.nr_address))

        # Informed
        informed_group = QGroupBox("Have you informed this person?")
        informed_group.setStyleSheet("""
            QGroupBox {
                font-weight: 500;
                font-size: 12px;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding-top: 16px;
            }
        """)
        informed_layout = QVBoxLayout(informed_group)

        self.nr_informed_yes = QRadioButton("Yes - I have informed them")
        self.nr_informed_no = QRadioButton("No - I have not yet informed them")
        self.nr_informed_yes.setChecked(True)

        self.informed_btn_group = QButtonGroup()
        self.informed_btn_group.addButton(self.nr_informed_yes)
        self.informed_btn_group.addButton(self.nr_informed_no)

        informed_layout.addWidget(self.nr_informed_yes)
        informed_layout.addWidget(self.nr_informed_no)
        nr_details_layout.addWidget(informed_group)

        layout.addWidget(self.nr_details_widget)

        # Unknown NR section
        self.nr_unknown_widget = QWidget()
        nr_unknown_layout = QVBoxLayout(self.nr_unknown_widget)
        nr_unknown_layout.setContentsMargins(0, 8, 0, 0)

        unknown_group = QGroupBox("Select option (a) or (b)")
        unknown_group.setStyleSheet("""
            QGroupBox {
                font-weight: 500;
                font-size: 12px;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding-top: 16px;
            }
        """)
        unknown_opt_layout = QVBoxLayout(unknown_group)

        self.nr_unable = QRadioButton("(a) Unable to ascertain who is the NR")
        self.nr_none = QRadioButton("(b) Patient has no NR within meaning of the Act")
        self.nr_unable.setChecked(True)

        self.unknown_btn_group = QButtonGroup()
        self.unknown_btn_group.addButton(self.nr_unable)
        self.unknown_btn_group.addButton(self.nr_none)

        unknown_opt_layout.addWidget(self.nr_unable)
        unknown_opt_layout.addWidget(self.nr_none)
        nr_unknown_layout.addWidget(unknown_group)

        layout.addWidget(self.nr_unknown_widget)
        self.nr_unknown_widget.hide()

        # Connect visibility toggle
        self.nr_known.toggled.connect(self._toggle_nr_sections)

        self.form_layout.addWidget(frame)

    def _toggle_nr_sections(self, known: bool):
        self.nr_details_widget.setVisible(known)
        self.nr_unknown_widget.setVisible(not known)

    # ----------------------------------------------------------------
    # SECTION 6: Patient Interview
    # ----------------------------------------------------------------
    def _create_section_6_interview(self):
        frame = self._create_section_frame(6, "Patient Interview")
        layout = frame.layout()

        date_row = QHBoxLayout()
        date_row.setSpacing(12)
        date_lbl = QLabel("Date last seen:")
        date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        date_lbl.setFixedWidth(120)
        date_row.addWidget(date_lbl)

        self.last_seen_date = QDateEdit()
        self.last_seen_date.setCalendarPopup(True)
        self.last_seen_date.setDate(QDate.currentDate())
        self.last_seen_date.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QDateEdit::drop-down {
                border: none;
                width: 24px;
            }
        """)
        date_row.addWidget(self.last_seen_date, 1)
        layout.addLayout(date_row)

        info = QLabel("This must be within 14 days ending on the day this application is completed.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 12px; color: #6b7280; padding: 8px; background: #f3f4f6; border-radius: 6px;")
        layout.addWidget(info)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 7: Medical Recommendations
    # ----------------------------------------------------------------
    def _create_section_7_medical_recs(self):
        frame = self._create_section_frame(7, "Medical Recommendations")
        layout = frame.layout()

        info = QLabel("This application is founded on two medical recommendations in the prescribed form.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 13px; color: #374151; padding: 12px; background: #dbeafe; border-radius: 6px;")
        layout.addWidget(info)

        reason_lbl = QLabel("If neither medical practitioner had previous acquaintance with the patient, explain why:")
        reason_lbl.setWordWrap(True)
        reason_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151; margin-top: 8px;")
        layout.addWidget(reason_lbl)

        self.no_acquaintance_reason = QTextEdit()
        self.no_acquaintance_reason.setPlaceholderText("Enter explanation (leave blank if not applicable)")
        self.no_acquaintance_reason.setMinimumHeight(100)
        self.no_acquaintance_reason.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus {
                border-color: #2563eb;
            }
        """)
        layout.addWidget(self.no_acquaintance_reason)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # SECTION 8: Signature
    # ----------------------------------------------------------------
    def _create_section_8_signature(self):
        frame = self._create_section_frame(8, "Signature")
        layout = frame.layout()

        date_row = QHBoxLayout()
        date_row.setSpacing(12)
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        date_lbl.setFixedWidth(120)
        date_row.addWidget(date_lbl)

        self.signature_date = QDateEdit()
        self.signature_date.setCalendarPopup(True)
        self.signature_date.setDate(QDate.currentDate())
        self.signature_date.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QDateEdit::drop-down {
                border: none;
                width: 24px;
            }
        """)
        date_row.addWidget(self.signature_date, 1)
        layout.addLayout(date_row)

        info = QLabel("The form will be signed manually after printing.")
        info.setStyleSheet("font-size: 12px; color: #6b7280; padding: 8px; background: #f3f4f6; border-radius: 6px;")
        layout.addWidget(info)

        self.form_layout.addWidget(frame)

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _go_back(self):
        self.go_back.emit()

    def _clear_form(self):
        reply = QMessageBox.question(
            self,
            "Clear Form",
            "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            # Clear all fields
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.local_authority.clear()
            self.approved_by_same.setChecked(True)
            self.approved_by_authority.clear()
            self.nr_known.setChecked(True)
            self.nr_option_a.setChecked(True)
            self.nr_name.clear()
            self.nr_address.clear()
            self.nr_informed_yes.setChecked(True)
            self.nr_unable.setChecked(True)
            self.last_seen_date.setDate(QDate.currentDate())
            self.no_acquaintance_reason.clear()
            self.signature_date.setDate(QDate.currentDate())

    def _import_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Import File",
            "",
            "All Files (*);;Word Documents (*.docx);;Text Files (*.txt)"
        )
        if file_path:
            QMessageBox.information(self, "Import", f"File selected: {file_path}\n\nData extraction coming soon.")

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Form A2",
            f"Form_A2_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            # Path to the template
            template_path = os.path.join(os.path.dirname(__file__), 'templates', 'Form_A2_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form A2 template not found.")
                return

            # Open the template
            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    run = para.add_run(new_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Data goes into WHITESPACE paragraphs below instruction paragraphs
            # Para 3: Hospital data (below para 2 instruction)
            # Para 5: AMHP data (below para 4 instruction)
            # Para 7: Patient data (below para 6 instruction)
            # Para 10: Local authority (below para 9 instruction)
            # Para 14: Approved authority (below para 13 instruction)
            # Para 18: NR option a (below para 17 instruction)
            # Para 21: NR option b (below para 20 instruction)
            # Para 29: Date seen (below para 28 instruction)
            # Para 34: No acquaintance reason (below para 33 instruction)

            # Get data from form fields
            hospital_text = self.hospital_name.text()
            if self.hospital_address.toPlainText():
                hospital_text += ", " + self.hospital_address.toPlainText().replace("\n", ", ")

            amhp_text = self.amhp_name.text()
            if self.amhp_address.toPlainText():
                amhp_text += ", " + self.amhp_address.toPlainText().replace("\n", ", ")
            if self.amhp_email.text():
                amhp_text += ", Email: " + self.amhp_email.text()

            patient_text = self.patient_name.text()
            if self.patient_address.toPlainText():
                patient_text += ", " + self.patient_address.toPlainText().replace("\n", ", ")

            # Fill Hospital - goes in para 3 (whitespace below instruction)
            if hospital_text.strip():
                set_para_text(paragraphs[3], hospital_text)
                highlight_yellow(paragraphs[3])

            # Fill AMHP - goes in para 5 (whitespace below instruction)
            if amhp_text.strip():
                set_para_text(paragraphs[5], amhp_text)
                highlight_yellow(paragraphs[5])

            # Fill Patient - goes in para 7 (whitespace below instruction)
            if patient_text.strip():
                set_para_text(paragraphs[7], patient_text)
                highlight_yellow(paragraphs[7])

            # Fill Local authority - goes in para 10 (whitespace below instruction)
            if self.local_authority.text():
                set_para_text(paragraphs[10], self.local_authority.text())
                highlight_yellow(paragraphs[10])

            # Handle approved by
            if self.approved_by_same.isChecked():
                for run in paragraphs[12].runs:
                    run.bold = True
                strikethrough_para(paragraphs[13])
                strikethrough_para(paragraphs[14])
            else:
                strikethrough_para(paragraphs[12])
                if self.approved_by_authority.text():
                    set_para_text(paragraphs[14], self.approved_by_authority.text())
                    highlight_yellow(paragraphs[14])

            # Fill NR section
            if self.nr_known.isChecked():
                nr_text = self.nr_name.text()
                if self.nr_address.toPlainText():
                    nr_text += ", " + self.nr_address.toPlainText().replace("\n", ", ")

                if self.nr_option_a.isChecked():
                    # Option a - NR is nearest relative - data in para 18
                    if nr_text.strip():
                        set_para_text(paragraphs[18], nr_text)
                        highlight_yellow(paragraphs[18])
                    strikethrough_para(paragraphs[20])
                    strikethrough_para(paragraphs[22])
                else:
                    # Option b - person exercising functions - data in para 21
                    if nr_text.strip():
                        set_para_text(paragraphs[21], nr_text)
                        highlight_yellow(paragraphs[21])
                    strikethrough_para(paragraphs[17])
                    strikethrough_para(paragraphs[18])
                    strikethrough_para(paragraphs[19])

                # Strikethrough unknown NR section
                strikethrough_para(paragraphs[24])
                strikethrough_para(paragraphs[25])
                strikethrough_para(paragraphs[26])

                # Informed status (para 23)
                para = paragraphs[23]
                for run in para.runs:
                    run.text = ""

                if self.nr_informed_yes.isChecked():
                    run1 = para.add_run("I have")
                    run1.font.name = 'Arial'
                    run1.font.size = Pt(12)

                    run2 = para.add_run("/have not yet*")
                    run2.font.name = 'Arial'
                    run2.font.size = Pt(12)
                    run2.font.strike = True

                    run3 = para.add_run(" informed that person that this application is to be made and of the nearest relative's power to order the discharge of the patient.")
                    run3.font.name = 'Arial'
                    run3.font.size = Pt(12)
                else:
                    run1 = para.add_run("I have/")
                    run1.font.name = 'Arial'
                    run1.font.size = Pt(12)
                    run1.font.strike = True

                    run2 = para.add_run("have not yet")
                    run2.font.name = 'Arial'
                    run2.font.size = Pt(12)

                    run3 = para.add_run("*")
                    run3.font.name = 'Arial'
                    run3.font.size = Pt(12)
                    run3.font.strike = True

                    run4 = para.add_run(" informed that person that this application is to be made and of the nearest relative's power to order the discharge of the patient.")
                    run4.font.name = 'Arial'
                    run4.font.size = Pt(12)
            else:
                # NR unknown - strikethrough known NR section (para 15-23)
                for i in range(15, 24):
                    strikethrough_para(paragraphs[i])

                if self.nr_unable.isChecked():
                    # Unable to ascertain - strikethrough option b
                    strikethrough_para(paragraphs[26])
                else:
                    # No NR - strikethrough option a
                    strikethrough_para(paragraphs[25])

            # Date last seen - goes in para 29 (whitespace below instruction)
            last_seen = self.last_seen_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[29], last_seen)
            highlight_yellow(paragraphs[29])

            # No acquaintance reason - goes in para 34 (whitespace below instruction)
            if self.no_acquaintance_reason.toPlainText():
                set_para_text(paragraphs[34], self.no_acquaintance_reason.toPlainText())
                highlight_yellow(paragraphs[34])

            # Signature date
            sig_date = self.signature_date.date().toString("dd MMMM yyyy")
            new_sig_text = f"Signed                                                                Date {sig_date}"
            set_para_text(paragraphs[36], new_sig_text)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form A2 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        """Get current form state for saving."""
        return {
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.toPlainText(),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.toPlainText(),
            "amhp_email": self.amhp_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.toPlainText(),
            "local_authority": self.local_authority.text(),
            "approved_by_same": self.approved_by_same.isChecked(),
            "approved_by_authority": self.approved_by_authority.text(),
            "nr_known": self.nr_known.isChecked(),
            "nr_option_a": self.nr_option_a.isChecked(),
            "nr_name": self.nr_name.text(),
            "nr_address": self.nr_address.toPlainText(),
            "nr_informed_yes": self.nr_informed_yes.isChecked(),
            "nr_unable": self.nr_unable.isChecked(),
            "last_seen_date": self.last_seen_date.date().toString("yyyy-MM-dd"),
            "no_acquaintance_reason": self.no_acquaintance_reason.toPlainText(),
            "signature_date": self.signature_date.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        """Load form state."""
        if not state:
            return
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setPlainText(state.get("hospital_address", ""))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setPlainText(state.get("amhp_address", ""))
        self.amhp_email.setText(state.get("amhp_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setPlainText(state.get("patient_address", ""))
        self.local_authority.setText(state.get("local_authority", ""))
        if state.get("approved_by_same", True):
            self.approved_by_same.setChecked(True)
        else:
            self.approved_by_different.setChecked(True)
        self.approved_by_authority.setText(state.get("approved_by_authority", ""))
        if state.get("nr_known", True):
            self.nr_known.setChecked(True)
        else:
            self.nr_unknown.setChecked(True)
        if state.get("nr_option_a", True):
            self.nr_option_a.setChecked(True)
        else:
            self.nr_option_b.setChecked(True)
        self.nr_name.setText(state.get("nr_name", ""))
        self.nr_address.setPlainText(state.get("nr_address", ""))
        if state.get("nr_informed_yes", True):
            self.nr_informed_yes.setChecked(True)
        else:
            self.nr_informed_no.setChecked(True)
        if state.get("nr_unable", True):
            self.nr_unable.setChecked(True)
        else:
            self.nr_none.setChecked(True)
        if state.get("last_seen_date"):
            self.last_seen_date.setDate(QDate.fromString(state["last_seen_date"], "yyyy-MM-dd"))
        self.no_acquaintance_reason.setPlainText(state.get("no_acquaintance_reason", ""))
        if state.get("signature_date"):
            self.signature_date.setDate(QDate.fromString(state["signature_date"], "yyyy-MM-dd"))
