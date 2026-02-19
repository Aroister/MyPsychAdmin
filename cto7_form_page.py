# ================================================================
#  CTO7 FORM PAGE — Report Extending Community Treatment Period
#  Mental Health Act 1983 - Form CTO7 Regulation 13(6)(a) and (b), 13(7)
#  Section 20A — Community treatment order: report extending
#  CARD/POPUP LAYOUT with ResizableSection
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QRadioButton, QButtonGroup, QComboBox, QSpinBox, QCompleter,
    QStyleFactory, QSlider, QStackedWidget, QSplitter, QSizePolicy
)
from utils.resource_path import resource_path
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# RESIZABLE SECTION WITH DRAG BAR
# ================================================================
class ResizableSection(QFrame):
    """Section with a draggable bottom edge to resize height."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._min_height = 120
        self._max_height = 600
        self._content = None
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0

        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(0)

        # Content container
        self._content_container = QWidget()
        self._content_layout = QVBoxLayout(self._content_container)
        self._content_layout.setContentsMargins(12, 8, 12, 8)
        self._layout.addWidget(self._content_container, 1)

        # Drag handle at bottom
        self._drag_handle = QFrame()
        self._drag_handle.setFixedHeight(8)
        self._drag_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        self._drag_handle.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
                border-radius: 2px;
            }
            QFrame:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.3 #2563eb, stop:0.7 #2563eb, stop:1 transparent);
            }
        """)
        self._drag_handle.mousePressEvent = self._handle_press
        self._drag_handle.mouseMoveEvent = self._handle_move
        self._drag_handle.mouseReleaseEvent = self._handle_release
        self._layout.addWidget(self._drag_handle)

    def set_content(self, widget):
        self._content = widget
        self._content_layout.addWidget(widget)

    def set_content_height(self, h):
        h = max(self._min_height, min(self._max_height, h))
        self._content_container.setFixedHeight(h)

    def _handle_press(self, event):
        self._dragging = True
        self._drag_start_y = event.globalPosition().y()
        self._drag_start_height = self._content_container.height()

    def _handle_move(self, event):
        if self._dragging:
            delta = event.globalPosition().y() - self._drag_start_y
            new_height = self._drag_start_height + delta
            new_height = max(self._min_height, min(self._max_height, new_height))
            self._content_container.setFixedHeight(int(new_height))

    def _handle_release(self, event):
        self._dragging = False


# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []

from shared_widgets import create_zoom_row


# ================================================================
# NO-WHEEL SLIDER
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# NO-WHEEL COMBOBOX
# ================================================================
class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelSpinBox(QSpinBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# CTO7 CARD WIDGET
# ================================================================
class CTO7CardWidget(QFrame):
    """Clickable card with editable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("cto7Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)

        self.setStyleSheet("""
            QFrame#cto7Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#cto7Card:hover {
                border-color: #2563eb;
                background: #eff6ff;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(6)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 19px; font-weight: 700; color: #2563eb;")
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 19px;
                color: #374151;
                background: transparent;
                border: none;
                padding: 0;
            }
        """)
        self.content.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        layout.addWidget(self.content, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=17)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_active(self, active: bool):
        """Set active state - shows persistent highlight."""
        if active:
            self.setStyleSheet("""
                QFrame#cto7Card {
                    background: #eff6ff;
                    border: 2px solid #2563eb;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#cto7Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#cto7Card:hover {
                    border-color: #2563eb;
                    background: #eff6ff;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# CTO7 TOOLBAR
# ================================================================
class CTO7Toolbar(QWidget):
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            CTO7Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN CTO7 FORM PAGE - Card/Popup Layout
# ================================================================
class CTO7FormPage(QWidget):
    """Page for completing MHA Form CTO7 - Report Extending CTO Period."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race", "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}

        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {"full_name": details[1] or "", "email": details[7] or ""}

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(50)
        header.setStyleSheet("background: #2563eb; border-bottom: 1px solid #1d4ed8;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 5px;
                font-size: 18px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO7 — Report Extending CTO Period")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - cards on left, popup on right with splitter
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(12, 12, 12, 12)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #2563eb, stop:0.7 #2563eb, stop:1 transparent);
            }
        """)

        # Left: Cards column (scrollable)
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(16, 16, 16, 16)
        self.cards_layout.setSpacing(8)

        # Create cards
        self._create_details_card()
        self._create_grounds_card()
        self._create_amhp_card()
        self._create_signatures_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        self.main_splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 10px;
            }
        """)

        # Create popup panels
        self._create_details_popup()
        self._create_grounds_popup()
        self._create_amhp_popup()
        self._create_signatures_popup()

        # Initialize cards with default date values
        self._update_details_card()
        self._update_amhp_card()
        self._update_signatures_card()

        self.main_splitter.addWidget(self.popup_stack)
        self.main_splitter.setSizes([280, 600])
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        # Connect live preview updates
        self._connect_grounds_live_preview()

        # Show first popup by default
        self._on_card_clicked("details")

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"details": 0, "grounds": 1, "amhp": 2, "signatures": 3}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            for k, card in self.cards.items():
                card.set_active(k == key)

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_details_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 350
        card = CTO7CardWidget("Details", "details")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["details"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_grounds_card(self):
        section = ResizableSection()
        section.set_content_height(200)
        section._min_height = 120
        section._max_height = 400
        card = CTO7CardWidget("Grounds for Extension", "grounds")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["grounds"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_amhp_card(self):
        section = ResizableSection()
        section.set_content_height(170)
        section._min_height = 120
        section._max_height = 350
        card = CTO7CardWidget("AMHP Statement", "amhp")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["amhp"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signatures_card(self):
        section = ResizableSection()
        section.set_content_height(150)
        section._min_height = 100
        section._max_height = 300
        card = CTO7CardWidget("Signatures", "signatures")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signatures"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_details_popup(self):
        """Combined popup for patient, hospital, RC details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        # Header
        header = QLabel("Patient & RC Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #2563eb;")
        popup_layout.addWidget(header)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        form = QWidget()
        form.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(form)
        form_layout.setContentsMargins(0, 0, 8, 0)
        form_layout.setSpacing(10)

        # Patient section
        patient_header = QLabel("Patient")
        patient_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 4px;")
        form_layout.addWidget(patient_header)

        self.patient_name = self._create_line_edit("Patient full name")
        form_layout.addWidget(self.patient_name)

        self.patient_address = self._create_line_edit("Patient address")
        form_layout.addWidget(self.patient_address)

        # Demographics row (age, gender, ethnicity - visible for clinical reasons)
        demo_row = QHBoxLayout()
        demo_row.setSpacing(8)

        # Age label and spin
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        demo_row.addWidget(age_lbl)
        self.age_spin = NoWheelSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setFixedWidth(60)
        self.age_spin.setStyleSheet("QSpinBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 18px; }")
        demo_row.addWidget(self.age_spin)

        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("O")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 17px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
        self.gender_group.addButton(self.gender_male, 0)
        self.gender_group.addButton(self.gender_female, 1)
        self.gender_group.addButton(self.gender_other, 2)
        demo_row.addWidget(self.gender_male)
        demo_row.addWidget(self.gender_female)
        demo_row.addWidget(self.gender_other)

        # Ethnicity combo - visible for clinical reasons
        self.ethnicity_combo = NoWheelComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(130)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 18px; }")
        demo_row.addWidget(self.ethnicity_combo)
        demo_row.addStretch()
        form_layout.addLayout(demo_row)

        # Hospital section
        hosp_header = QLabel("Hospital")
        hosp_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(hosp_header)

        self.hospital = self._create_line_edit("Hospital name and address")
        form_layout.addWidget(self.hospital)

        # RC section
        rc_header = QLabel("Responsible Clinician")
        rc_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(rc_header)

        self.rc_name = self._create_line_edit("RC full name")
        form_layout.addWidget(self.rc_name)

        rc_row = QHBoxLayout()
        rc_row.setSpacing(8)
        self.rc_address = self._create_line_edit("Address")
        rc_row.addWidget(self.rc_address, 2)
        self.rc_email = self._create_line_edit("Email")
        rc_row.addWidget(self.rc_email, 1)
        form_layout.addLayout(rc_row)

        # Dates
        dates_header = QLabel("CTO Dates")
        dates_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(dates_header)

        dates_row = QHBoxLayout()
        dates_row.setSpacing(12)
        cto_lbl = QLabel("CTO made:")
        cto_lbl.setStyleSheet("font-size: 17px; color: #6b7280;")
        dates_row.addWidget(cto_lbl)
        self.cto_date = self._create_date_edit()
        self.cto_date.setFixedWidth(120)
        dates_row.addWidget(self.cto_date)
        exam_lbl = QLabel("Examined:")
        exam_lbl.setStyleSheet("font-size: 17px; color: #6b7280;")
        dates_row.addWidget(exam_lbl)
        self.exam_date = self._create_date_edit()
        self.exam_date.setFixedWidth(120)
        dates_row.addWidget(self.exam_date)
        dates_row.addStretch()
        form_layout.addLayout(dates_row)

        form_layout.addStretch()
        scroll.setWidget(form)
        popup_layout.addWidget(scroll, 1)

        # Auto-sync to card
        self.patient_name.textChanged.connect(self._update_details_card)
        self.patient_address.textChanged.connect(self._update_details_card)
        self.hospital.textChanged.connect(self._update_details_card)
        self.rc_name.textChanged.connect(self._update_details_card)
        self.rc_address.textChanged.connect(self._update_details_card)
        self.exam_date.dateChanged.connect(self._update_details_card)

        self.popup_stack.addWidget(popup)

    def _create_grounds_popup(self):
        """Popup for grounds with live preview."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(0, 0, 0, 0)
        popup_layout.setSpacing(0)

        # Hidden label for state storage (preview removed - card auto-syncs)
        self.grounds_preview = QLabel("")
        self.grounds_preview.hide()

        # Scrollable form below
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("QScrollArea { background: white; border: none; }")

        form_container = QWidget()
        form_container.setStyleSheet("background: white;")
        form_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(12, 8, 12, 12)
        form_layout.setSpacing(8)

        # Mental Disorder section
        md_frame = QFrame()
        md_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 6px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(10, 8, 10, 8)
        md_layout.setSpacing(4)
        md_header = QLabel("Mental Disorder (ICD-10)")
        md_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        # Primary diagnosis dropdown with grouped items
        self.dx_primary = NoWheelComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select primary diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_primary.currentTextChanged.connect(self._update_grounds_preview)
        md_layout.addWidget(self.dx_primary)

        # Secondary diagnosis dropdown
        self.dx_secondary = NoWheelComboBox()
        self.dx_secondary.setEditable(True)
        self.dx_secondary.lineEdit().setPlaceholderText("Secondary diagnosis (optional)...")
        self.dx_secondary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_secondary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_secondary.addItem(f"── {group_name} ──", None)
            idx = self.dx_secondary.count() - 1
            self.dx_secondary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_secondary.addItem(dx, dx)
        completer2 = QCompleter(ICD10_FLAT, self.dx_secondary)
        completer2.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer2.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_secondary.setCompleter(completer2)
        self.dx_secondary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_secondary.currentTextChanged.connect(self._update_grounds_preview)
        md_layout.addWidget(self.dx_secondary)

        form_layout.addWidget(md_frame)

        # Legal Criteria section
        lc_frame = QFrame()
        lc_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 6px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(10, 8, 10, 8)
        lc_layout.setSpacing(4)
        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)
        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._update_grounds_preview)
        nature_opt_layout.addWidget(self.relapsing_cb)
        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._update_grounds_preview)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)
        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._update_grounds_preview)
        nature_opt_layout.addWidget(self.chronic_cb)
        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)
        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(100)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)
        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)
        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_grounds_preview)
        degree_opt_layout.addWidget(self.degree_details)
        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity section
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)
        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)
        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)
        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._update_grounds_preview)
        mh_opt_layout.addWidget(self.poor_compliance_cb)
        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._update_grounds_preview)
        mh_opt_layout.addWidget(self.limited_insight_cb)
        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)
        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)
        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._update_grounds_preview)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)
        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)
        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)
        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)
        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_hist_neglect)
        self.self_hist_risky = QCheckBox("Risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_hist_risky)
        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_hist_harm)
        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        self_opt_layout.addWidget(self_curr_lbl)
        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_curr_neglect)
        self.self_curr_risky = QCheckBox("Risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_curr_risky)
        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_curr_harm)
        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)
        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)
        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)
        self.others_hist_violence = QCheckBox("Violence")
        self.others_hist_violence.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_hist_violence)
        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_hist_verbal)
        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        others_opt_layout.addWidget(others_curr_lbl)
        self.others_curr_violence = QCheckBox("Violence")
        self.others_curr_violence.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_curr_violence)
        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 11px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_curr_verbal)
        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)
        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)
        form_layout.addWidget(lc_frame)

        # Extension section
        ext_frame = QFrame()
        ext_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        ext_frame.setStyleSheet("QFrame { background: #fef2f2; border: none; border-radius: 6px; }")
        ext_layout = QVBoxLayout(ext_frame)
        ext_layout.setContentsMargins(10, 8, 10, 8)
        ext_layout.setSpacing(4)
        ext_header = QLabel("Extension Because")
        ext_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #991b1b;")
        ext_layout.addWidget(ext_header)
        self.cto_effective_cb = QCheckBox("CTO effective")
        self.cto_effective_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.cto_effective_cb.toggled.connect(self._update_grounds_preview)
        ext_layout.addWidget(self.cto_effective_cb)
        self.conditions_met_cb = QCheckBox("Conditions continue to be met")
        self.conditions_met_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.conditions_met_cb.toggled.connect(self._update_grounds_preview)
        ext_layout.addWidget(self.conditions_met_cb)
        self.needs_supervision_cb = QCheckBox("Needs continued supervision")
        self.needs_supervision_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.needs_supervision_cb.toggled.connect(self._update_grounds_preview)
        ext_layout.addWidget(self.needs_supervision_cb)
        form_layout.addWidget(ext_frame)

        form_layout.addStretch()
        form_scroll.setWidget(form_container)
        popup_layout.addWidget(form_scroll, 1)

        self.popup_stack.addWidget(popup)

    def _create_amhp_popup(self):
        """Popup for AMHP details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("AMHP Statement (Part 2)")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #059669;")
        popup_layout.addWidget(header)

        # Form
        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        self.amhp_name = self._create_line_edit("AMHP full name")
        form_layout.addWidget(self.amhp_name)

        self.amhp_address = self._create_line_edit("AMHP address")
        form_layout.addWidget(self.amhp_address)

        self.amhp_email = self._create_line_edit("Email")
        form_layout.addWidget(self.amhp_email)

        auth_row = QHBoxLayout()
        auth_row.setSpacing(8)
        self.amhp_authority = self._create_line_edit("Acting on behalf of (Authority)")
        auth_row.addWidget(self.amhp_authority, 1)
        self.amhp_approved_by = self._create_line_edit("Approved by (if different)")
        auth_row.addWidget(self.amhp_approved_by, 1)
        form_layout.addLayout(auth_row)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(8)
        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 17px; color: #6b7280;")
        sig_row.addWidget(sig_lbl)
        self.amhp_sig_date = self._create_date_edit()
        self.amhp_sig_date.setFixedWidth(120)
        sig_row.addWidget(self.amhp_sig_date)
        sig_row.addStretch()
        form_layout.addLayout(sig_row)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        # Auto-sync to card
        self.amhp_name.textChanged.connect(self._update_amhp_card)
        self.amhp_address.textChanged.connect(self._update_amhp_card)
        self.amhp_authority.textChanged.connect(self._update_amhp_card)
        self.amhp_sig_date.dateChanged.connect(self._update_amhp_card)

        self.popup_stack.addWidget(popup)

    def _create_signatures_popup(self):
        """Popup for signature and consulted person details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Signatures & Consultation (Part 3)")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #7c3aed;")
        popup_layout.addWidget(header)

        # Form
        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        consult_lbl = QLabel("I consulted:")
        consult_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        form_layout.addWidget(consult_lbl)

        self.consulted_name = self._create_line_edit("Full name")
        form_layout.addWidget(self.consulted_name)

        self.consulted_profession = NoWheelComboBox()
        self.consulted_profession.addItem("Select profession...")
        self.consulted_profession.addItems([
            "Registered Mental Health Nurse",
            "Registered Learning Disabilities Nurse",
            "Occupational Therapist",
            "Social Worker",
            "Psychologist"
        ])
        self.consulted_profession.setStyleSheet("QComboBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 5px; font-size: 18px; }")
        form_layout.addWidget(self.consulted_profession)

        # RC Signature Date
        rc_sig_lbl = QLabel("RC Signature Date:")
        rc_sig_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(rc_sig_lbl)

        rc_sig_row = QHBoxLayout()
        rc_sig_row.setSpacing(8)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(120)
        rc_sig_row.addWidget(self.rc_sig_date)
        rc_sig_row.addStretch()
        form_layout.addLayout(rc_sig_row)

        # Final Signature Date
        final_lbl = QLabel("Final Signature Date:")
        final_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(final_lbl)

        final_row = QHBoxLayout()
        final_row.setSpacing(8)
        self.final_sig_date = self._create_date_edit()
        self.final_sig_date.setFixedWidth(120)
        final_row.addWidget(self.final_sig_date)
        final_row.addStretch()
        form_layout.addLayout(final_row)

        # Furnishing report method
        furnish_lbl = QLabel("I am furnishing this report by:")
        furnish_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(furnish_lbl)

        self.furnish_group = QButtonGroup(self)
        self.furnish_internal = QRadioButton("Internal mail system")
        self.furnish_electronic = QRadioButton("Electronic communication")
        self.furnish_other = QRadioButton("Other delivery method")
        self.furnish_internal.setChecked(True)
        for rb in [self.furnish_internal, self.furnish_electronic, self.furnish_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 17px;
                    color: #374151;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
            self.furnish_group.addButton(rb)
            form_layout.addWidget(rb)

        warning = QLabel("This report is NOT VALID unless Parts 1, 2 & 3 are completed and signed.")
        warning.setWordWrap(True)
        warning.setStyleSheet("font-size: 17px; color: #dc2626; font-weight: 600; padding: 8px; background: #fef2f2; border-radius: 6px; margin-top: 12px;")
        form_layout.addWidget(warning)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        # Auto-sync to card
        self.rc_sig_date.dateChanged.connect(self._update_signatures_card)
        self.final_sig_date.dateChanged.connect(self._update_signatures_card)

        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_details_card(self):
        parts = []
        if self.patient_name.text().strip():
            parts.append(self.patient_name.text().strip())
        if self.patient_address.text().strip():
            parts.append(self.patient_address.text().strip())
        if self.hospital.text().strip():
            parts.append(self.hospital.text().strip())
        if self.rc_name.text().strip():
            parts.append(self.rc_name.text().strip())
        if self.rc_address.text().strip():
            parts.append(self.rc_address.text().strip())
        self.cards["details"].set_content_text("\n".join(parts))

    def _update_grounds_card(self):
        self.cards["grounds"].set_content_text(self.grounds_preview.text())

    def _update_amhp_card(self):
        parts = []
        if self.amhp_name.text().strip():
            parts.append(self.amhp_name.text().strip())
        if self.amhp_address.text().strip():
            parts.append(self.amhp_address.text().strip())
        if self.amhp_authority.text().strip():
            parts.append(self.amhp_authority.text().strip())
        if self.amhp_sig_date.date().isValid():
            parts.append(self.amhp_sig_date.date().toString('dd MMM yyyy'))
        self.cards["amhp"].set_content_text("\n".join(parts))

    def _update_signatures_card(self):
        parts = []
        if self.consulted_name.text().strip():
            parts.append(self.consulted_name.text().strip())
        if self.rc_sig_date.date().isValid():
            parts.append(self.rc_sig_date.date().toString('dd MMM yyyy'))
        if self.final_sig_date.date().isValid():
            parts.append(self.final_sig_date.date().toString('dd MMM yyyy'))
        self.cards["signatures"].set_content_text("\n".join(parts))

    # ----------------------------------------------------------------
    # HELPERS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("QLineEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 8px; font-size: 18px; } QLineEdit:focus { border-color: #2563eb; }")
        return edit

    def _create_date_edit(self) -> NoWheelDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 6px; font-size: 18px; }")
        return date_edit

    # ----------------------------------------------------------------
    # TOGGLE HANDLERS
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_grounds_preview()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_grounds_preview()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_grounds_preview()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm, self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
                cb.setChecked(False)
        self._update_grounds_preview()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            for cb in [self.others_hist_violence, self.others_hist_verbal, self.others_curr_violence, self.others_curr_verbal]:
                cb.setChecked(False)
        self._update_grounds_preview()

    # ----------------------------------------------------------------
    # LIVE PREVIEW
    # ----------------------------------------------------------------
    def _connect_grounds_live_preview(self):
        self.patient_name.textChanged.connect(self._update_grounds_preview)
        self.gender_male.toggled.connect(self._update_grounds_preview)
        self.gender_female.toggled.connect(self._update_grounds_preview)
        self.gender_other.toggled.connect(self._update_grounds_preview)

    def _update_grounds_preview(self):
        text = self._generate_grounds_text()
        self.grounds_preview.setText(text)
        self._update_grounds_card()

    def _generate_grounds_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"
        paragraphs = []

        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        if self.dx_primary.currentText().strip():
            diagnoses.append(self.dx_primary.currentText().strip())
        if self.dx_secondary.currentText().strip():
            diagnoses.append(self.dx_secondary.currentText().strip())

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for the patient to receive medical treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for the patient to receive medical treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for the patient to receive medical treatment.")
            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")
            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")
        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        para2_parts = []
        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("protection of others")
        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"The CTO continues to be necessary for {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"The CTO continues to be necessary for {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"The CTO continues to be necessary for {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")
        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_items = []
            if self.poor_compliance_cb.isChecked():
                mh_items.append("poor compliance with treatment")
            if self.limited_insight_cb.isChecked():
                mh_items.append(f"limited insight into {p['pos_l']} illness")
            if mh_items:
                para2_parts.append(f"Regarding {p['pos_l']} health, I would be concerned about {' and '.join(mh_items)} resulting in deterioration if not on a CTO.")
        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health: {details}.")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            if self.gender_male.isChecked():
                reflexive = "himself"
            elif self.gender_female.isChecked():
                reflexive = "herself"
            else:
                reflexive = "themselves"
            risk_types = [("self neglect", self.self_hist_neglect.isChecked(), self.self_curr_neglect.isChecked()), (f"placing of {reflexive} in risky situations", self.self_hist_risky.isChecked(), self.self_curr_risky.isChecked()), ("self harm", self.self_hist_harm.isChecked(), self.self_curr_harm.isChecked())]
            both_items, hist_only, curr_only = [], [], []
            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)
            self_text = f"With respect to {p['pos_l']} own safety, if not on a CTO, I would be concerned about"
            parts = []
            if both_items:
                parts.append(f"historical and current {', '.join(both_items[:-1])}, and of {both_items[-1]}" if len(both_items) > 1 else f"historical and current {both_items[0]}")
            if hist_only:
                parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}" if len(hist_only) > 1 else f"historical {hist_only[0]}")
            if curr_only:
                parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}" if len(curr_only) > 1 else f"current {curr_only[0]}")
            if parts:
                self_text += " " + ", and ".join(parts) + "."
                para2_parts.append(self_text)
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            risk_types = [("violence to others", self.others_hist_violence.isChecked(), self.others_curr_violence.isChecked()), ("verbal aggression", self.others_hist_verbal.isChecked(), self.others_curr_verbal.isChecked())]
            both_items, hist_only, curr_only = [], [], []
            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)
            others_text = f"Regarding risk to others, if not on a CTO, I would be concerned about the risk of"
            parts = []
            if both_items:
                parts.append(f"{both_items[0]} which is both historical and current" if len(both_items) == 1 else f"{', '.join(both_items[:-1])}, and {both_items[-1]} which are both historical and current")
            if hist_only:
                parts.append(f"historical {hist_only[0]}" if len(hist_only) == 1 else f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
            if curr_only:
                parts.append(f"current {curr_only[0]}" if len(curr_only) == 1 else f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")
            if parts:
                others_text += " " + " and of ".join(parts) + "."
                para2_parts.append(others_text)
        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        para3_parts = []
        ext_reasons = []
        if self.cto_effective_cb.isChecked():
            ext_reasons.append("the CTO has been effective in managing the patient's care")
        if self.conditions_met_cb.isChecked():
            ext_reasons.append("the conditions for the CTO continue to be met")
        if self.needs_supervision_cb.isChecked():
            ext_reasons.append(f"{p['subj_l']} requires continued supervision in the community")
        if ext_reasons:
            para3_parts.append(f"The CTO should be extended because {', and '.join(ext_reasons)}.")
        if para3_parts:
            paragraphs.append(" ".join(para3_parts))
        return "\n\n".join(paragraphs)

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.hospital.clear()
            self.rc_name.clear()
            self.rc_address.clear()
            self.rc_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.cto_date.setDate(QDate.currentDate())
            self.exam_date.setDate(QDate.currentDate())
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            self.cto_effective_cb.setChecked(False)
            self.conditions_met_cb.setChecked(False)
            self.needs_supervision_cb.setChecked(False)
            self.grounds_preview.setText("")
            self.rc_sig_date.setDate(QDate.currentDate())
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_email.clear()
            self.amhp_authority.clear()
            self.amhp_approved_by.clear()
            self.amhp_sig_date.setDate(QDate.currentDate())
            self.consulted_name.clear()
            self.consulted_profession.setCurrentIndex(0)
            self.final_sig_date.setDate(QDate.currentDate())
            for card in self.cards.values():
                card.set_content_text("")
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Form CTO7", f"Form_CTO7_{datetime.now().strftime('%Y%m%d')}.docx", "Word Documents (*.docx)")
        if not file_path:
            return
        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_CTO7_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO7 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color (#918C0D) and cream highlight (#FFFED5)
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)
            CREAM_FILL = 'FFFED5'

            # Pre-process: clean ALL paragraphs - remove permission markers, convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), CREAM_FILL)
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), CREAM_FILL)

            def set_entry_box(para, content: str):
                """Set entry box with gold brackets - content between brackets is cream."""
                if not content or not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Opening bracket
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def set_labeled_entry(para, label: str, content: str, suffix: str = ""):
                """Set labeled entry - only content between brackets is cream."""
                if not content or not content.strip():
                    content = '                                        '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Label
                if para.runs:
                    para.runs[0].text = label
                    label_run = para.runs[0]
                else:
                    label_run = para.add_run(label)
                label_run.font.name = 'Arial'
                label_run.font.size = Pt(12)
                # Opening bracket
                bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)
                # Suffix
                if suffix:
                    suffix_run = para.add_run(suffix)
                    suffix_run.font.name = 'Arial'
                    suffix_run.font.size = Pt(12)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            # Build data
            hospital_text = self.hospital.text().strip()
            rc_text = self.rc_name.text().strip()
            if self.rc_address.text().strip():
                rc_text += ", " + self.rc_address.text().strip()
            if self.rc_email.text().strip():
                rc_text += ", " + self.rc_email.text().strip()
            patient_text = self.patient_name.text().strip()
            if self.patient_address.text().strip():
                patient_text += ", " + self.patient_address.text().strip()
            cto_date = self.cto_date.date().toString("dd/MM/yyyy")
            exam_date = self.exam_date.date().toString("dd/MM/yyyy")
            grounds_text = self.cards["grounds"].get_content() if "grounds" in self.cards else ""
            rc_sig_date = self.rc_sig_date.date().toString("dd/MM/yyyy")
            amhp_text = self.amhp_name.text().strip()
            if self.amhp_address.text().strip():
                amhp_text += ", " + self.amhp_address.text().strip()
            if self.amhp_email.text().strip():
                amhp_text += ", " + self.amhp_email.text().strip()
            authority_text = self.amhp_authority.text().strip()
            approved_by_text = self.amhp_approved_by.text().strip() if hasattr(self, 'amhp_approved_by') else ""
            amhp_sig_date = self.amhp_sig_date.date().toString("dd/MM/yyyy")
            consult_text = self.consulted_name.text().strip()
            if hasattr(self, 'consulted_profession') and self.consulted_profession.currentIndex() > 0:
                consult_text += ", " + self.consulted_profession.currentText()
            final_sig_date = self.final_sig_date.date().toString("dd/MM/yyyy")

            # Strikethrough logic - only if at least one option selected
            any_necessity_selected = self.health_cb.isChecked() or self.safety_cb.isChecked() or self.others_cb.isChecked()
            any_furnish_selected = self.furnish_internal.isChecked() or self.furnish_electronic.isChecked() or self.furnish_other.isChecked()

            # Track filled boxes
            hospital_filled = False
            rc_filled = False
            patient_filled = False
            cto_date_filled = False
            exam_date_filled = False
            grounds_filled = False
            grounds_continuation_handled = False  # Grounds spans 2 paragraphs
            amhp_filled = False
            authority_filled = False
            approved_by_filled = False
            consult_filled = False
            part4_received_by_me_found = False  # Track when we hit Part 4 section
            part4_placeholder_filled = False  # Track Part 4 closing bracket placeholder

            # Debug: print template structure
            print("\n=== CTO7 Template Structure ===")
            for i, para in enumerate(doc.paragraphs):
                txt = para.text[:80].replace('\n', ' ') if para.text else "(empty)"
                print(f"DEBUG: Para {i}: {txt}")
            print("=== End Structure ===\n")

            # Process paragraphs
            for i, para in enumerate(doc.paragraphs):
                text = para.text
                text_stripped = text.strip()
                text_lower = text_stripped.lower()

                # (i), (ii), (iii) necessity options - Word numbering provides labels
                # (i) gets opening bracket, (ii) no brackets, (iii) gets closing bracket
                if ("health" in text_lower) and "patient" in text_lower and "safety" not in text_lower and "protection" not in text_lower and len(text_stripped) < 50:
                    strike_this = any_necessity_selected and not self.health_cb.isChecked()
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Opening bracket
                    if para.runs:
                        para.runs[0].text = '['
                        ob = para.runs[0]
                    else:
                        ob = para.add_run('[')
                    ob.font.name = 'Arial'
                    ob.font.size = Pt(12)
                    ob.font.bold = True
                    ob.font.color.rgb = BRACKET_COLOR
                    rPr_ob = ob._element.get_or_add_rPr()
                    shd_ob = OxmlElement('w:shd')
                    shd_ob.set(qn('w:val'), 'clear')
                    shd_ob.set(qn('w:color'), 'auto')
                    shd_ob.set(qn('w:fill'), CREAM_FILL)
                    rPr_ob.append(shd_ob)
                    if strike_this:
                        ob.font.strike = True
                    # Content
                    ct = para.add_run("the patient's health")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    if strike_this:
                        ct.font.strike = True
                    print(f"DEBUG: Formatted (i) health at para {i}")
                    continue

                if ("safety" in text_lower) and "patient" in text_lower and "health" not in text_lower and "protection" not in text_lower and len(text_stripped) < 50:
                    strike_this = any_necessity_selected and not self.safety_cb.isChecked()
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Content only - no brackets (middle option)
                    if para.runs:
                        para.runs[0].text = "the patient's safety"
                        ct = para.runs[0]
                    else:
                        ct = para.add_run("the patient's safety")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    if strike_this:
                        ct.font.strike = True
                    print(f"DEBUG: Formatted (ii) safety at para {i}")
                    continue

                if "protection" in text_lower and "other" in text_lower and "person" in text_lower and len(text_stripped) < 50:
                    strike_this = any_necessity_selected and not self.others_cb.isChecked()
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Content
                    if para.runs:
                        para.runs[0].text = "the protection of other persons"
                        ct = para.runs[0]
                    else:
                        ct = para.add_run("the protection of other persons")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    if strike_this:
                        ct.font.strike = True
                    # Closing bracket
                    cb = para.add_run(']')
                    cb.font.name = 'Arial'
                    cb.font.size = Pt(12)
                    cb.font.bold = True
                    cb.font.color.rgb = BRACKET_COLOR
                    rPr_cb = cb._element.get_or_add_rPr()
                    shd_cb = OxmlElement('w:shd')
                    shd_cb.set(qn('w:val'), 'clear')
                    shd_cb.set(qn('w:color'), 'auto')
                    shd_cb.set(qn('w:fill'), CREAM_FILL)
                    rPr_cb.append(shd_cb)
                    if strike_this:
                        cb.font.strike = True
                    print(f"DEBUG: Formatted (iii) protection at para {i}")
                    continue

                # Entry boxes - whitespace-only paragraphs (template has these as placeholders)
                # Skip empty strings but catch whitespace-only paragraphs
                is_placeholder = len(text) > 10 and (not text_stripped or text.isspace() or text_stripped.replace(' ', '') == '')
                if is_placeholder:
                    # Grounds section spans 2 paragraphs (para 27 and 28)
                    # After filling grounds, skip the continuation paragraph
                    if grounds_filled and not grounds_continuation_handled:
                        # This is the continuation of grounds - just add cream highlight, no brackets
                        for run in para.runs:
                            run.text = ""
                        while len(para.runs) > 1:
                            para._element.remove(para.runs[-1]._element)
                        pPr = para._element.get_or_add_pPr()
                        for old_shd in pPr.findall(qn('w:shd')):
                            pPr.remove(old_shd)
                        # Just add cream highlight with closing bracket
                        if para.runs:
                            para.runs[0].text = '                                                                                                    '
                            ct = para.runs[0]
                        else:
                            ct = para.add_run('                                                                                                    ')
                        ct.font.name = 'Arial'
                        ct.font.size = Pt(12)
                        rPr_ct = ct._element.get_or_add_rPr()
                        shd_ct = OxmlElement('w:shd')
                        shd_ct.set(qn('w:val'), 'clear')
                        shd_ct.set(qn('w:color'), 'auto')
                        shd_ct.set(qn('w:fill'), CREAM_FILL)
                        rPr_ct.append(shd_ct)
                        # Add closing bracket at end
                        cb = para.add_run(']')
                        cb.font.name = 'Arial'
                        cb.font.size = Pt(12)
                        cb.font.bold = True
                        cb.font.color.rgb = BRACKET_COLOR
                        rPr_cb = cb._element.get_or_add_rPr()
                        shd_cb = OxmlElement('w:shd')
                        shd_cb.set(qn('w:val'), 'clear')
                        shd_cb.set(qn('w:color'), 'auto')
                        shd_cb.set(qn('w:fill'), CREAM_FILL)
                        rPr_cb.append(shd_cb)
                        grounds_continuation_handled = True
                        print(f"DEBUG: Formatted grounds continuation at para {i}")
                        continue

                    if not hospital_filled:
                        set_entry_box(para, hospital_text)
                        hospital_filled = True
                        print(f"DEBUG: Filled hospital at para {i}")
                        continue
                    elif not rc_filled:
                        set_entry_box(para, rc_text)
                        rc_filled = True
                        print(f"DEBUG: Filled RC at para {i}")
                        continue
                    elif not patient_filled:
                        set_entry_box(para, patient_text)
                        patient_filled = True
                        print(f"DEBUG: Filled patient at para {i}")
                        continue
                    elif not cto_date_filled:
                        set_entry_box(para, cto_date)
                        cto_date_filled = True
                        print(f"DEBUG: Filled CTO date at para {i}")
                        continue
                    elif not exam_date_filled:
                        set_entry_box(para, exam_date)
                        exam_date_filled = True
                        print(f"DEBUG: Filled exam date at para {i}")
                        continue
                    elif not grounds_filled:
                        # Grounds has opening bracket only (closing is on continuation line)
                        for run in para.runs:
                            run.text = ""
                        while len(para.runs) > 1:
                            para._element.remove(para.runs[-1]._element)
                        pPr = para._element.get_or_add_pPr()
                        for old_shd in pPr.findall(qn('w:shd')):
                            pPr.remove(old_shd)
                        # Opening bracket
                        if para.runs:
                            para.runs[0].text = '['
                            bracket_open = para.runs[0]
                        else:
                            bracket_open = para.add_run('[')
                        bracket_open.font.name = 'Arial'
                        bracket_open.font.size = Pt(12)
                        bracket_open.font.bold = True
                        bracket_open.font.color.rgb = BRACKET_COLOR
                        rPr_ob = bracket_open._element.get_or_add_rPr()
                        shd_ob = OxmlElement('w:shd')
                        shd_ob.set(qn('w:val'), 'clear')
                        shd_ob.set(qn('w:color'), 'auto')
                        shd_ob.set(qn('w:fill'), CREAM_FILL)
                        rPr_ob.append(shd_ob)
                        # Content (no closing bracket - continues to next line)
                        content = grounds_text if grounds_text.strip() else '                                                                                                    '
                        content_run = para.add_run(content)
                        content_run.font.name = 'Arial'
                        content_run.font.size = Pt(12)
                        rPr2 = content_run._element.get_or_add_rPr()
                        shd2 = OxmlElement('w:shd')
                        shd2.set(qn('w:val'), 'clear')
                        shd2.set(qn('w:color'), 'auto')
                        shd2.set(qn('w:fill'), CREAM_FILL)
                        rPr2.append(shd2)
                        grounds_filled = True
                        print(f"DEBUG: Filled grounds (opening) at para {i}")
                        continue
                    elif not amhp_filled:
                        set_entry_box(para, amhp_text)
                        amhp_filled = True
                        print(f"DEBUG: Filled AMHP at para {i}")
                        continue
                    elif not authority_filled:
                        set_entry_box(para, authority_text)
                        authority_filled = True
                        print(f"DEBUG: Filled authority at para {i}")
                        continue
                    elif not approved_by_filled:
                        set_entry_box(para, approved_by_text)
                        approved_by_filled = True
                        print(f"DEBUG: Filled approved_by at para {i}")
                        continue
                    elif not consult_filled:
                        set_entry_box(para, consult_text)
                        consult_filled = True
                        print(f"DEBUG: Filled consult at para {i}")
                        continue

                # "that authority" with golden brackets
                if text_stripped.lower() == "that authority" or text_lower == "that authority":
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    if para.runs:
                        para.runs[0].text = '['
                        ob = para.runs[0]
                    else:
                        ob = para.add_run('[')
                    ob.font.name = 'Arial'
                    ob.font.size = Pt(12)
                    ob.font.bold = True
                    ob.font.color.rgb = BRACKET_COLOR
                    rPr_ob = ob._element.get_or_add_rPr()
                    shd_ob = OxmlElement('w:shd')
                    shd_ob.set(qn('w:val'), 'clear')
                    shd_ob.set(qn('w:color'), 'auto')
                    shd_ob.set(qn('w:fill'), CREAM_FILL)
                    rPr_ob.append(shd_ob)
                    ct = para.add_run("that authority")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    cb = para.add_run(']')
                    cb.font.name = 'Arial'
                    cb.font.size = Pt(12)
                    cb.font.bold = True
                    cb.font.color.rgb = BRACKET_COLOR
                    rPr_cb = cb._element.get_or_add_rPr()
                    shd_cb = OxmlElement('w:shd')
                    shd_cb.set(qn('w:val'), 'clear')
                    shd_cb.set(qn('w:color'), 'auto')
                    shd_cb.set(qn('w:fill'), CREAM_FILL)
                    rPr_cb.append(shd_cb)
                    print(f"DEBUG: Formatted 'that authority' at para {i}")
                    continue

                # "indicate here" checkbox
                if "indicate here" in text_lower and "attach" in text_lower:
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    if para.runs:
                        para.runs[0].text = '['
                        para.runs[0].font.name = 'Arial'
                        para.runs[0].font.size = Pt(12)
                    else:
                        r = para.add_run('[')
                        r.font.name = 'Arial'
                        r.font.size = Pt(12)
                    text1 = para.add_run('If you need to continue on a separate sheet please indicate here ')
                    text1.font.name = 'Arial'
                    text1.font.size = Pt(12)
                    cb_open = para.add_run('[')
                    cb_open.font.bold = True
                    cb_open.font.color.rgb = BRACKET_COLOR
                    rPr_cbo = cb_open._element.get_or_add_rPr()
                    shd_cbo = OxmlElement('w:shd')
                    shd_cbo.set(qn('w:val'), 'clear')
                    shd_cbo.set(qn('w:color'), 'auto')
                    shd_cbo.set(qn('w:fill'), CREAM_FILL)
                    rPr_cbo.append(shd_cbo)
                    cb_space = para.add_run(' ')
                    rPr_cbs = cb_space._element.get_or_add_rPr()
                    shd_cbs = OxmlElement('w:shd')
                    shd_cbs.set(qn('w:val'), 'clear')
                    shd_cbs.set(qn('w:color'), 'auto')
                    shd_cbs.set(qn('w:fill'), CREAM_FILL)
                    rPr_cbs.append(shd_cbs)
                    cb_close = para.add_run(']')
                    cb_close.font.bold = True
                    cb_close.font.color.rgb = BRACKET_COLOR
                    rPr_cbc = cb_close._element.get_or_add_rPr()
                    shd_cbc = OxmlElement('w:shd')
                    shd_cbc.set(qn('w:val'), 'clear')
                    shd_cbc.set(qn('w:color'), 'auto')
                    shd_cbc.set(qn('w:fill'), CREAM_FILL)
                    rPr_cbc.append(shd_cbc)
                    text2 = para.add_run(' and attach that sheet to this form]')
                    text2.font.name = 'Arial'
                    text2.font.size = Pt(12)
                    print(f"DEBUG: Formatted 'indicate here' checkbox at para {i}")
                    continue

                # Signed lines - use text content to identify which Part
                if text_stripped.lower().startswith("signed"):
                    # Part 4: "on behalf of the managers" - appears after Part 4 placeholders
                    if "on behalf of the managers" in text_lower or (part4_received_by_me_found and part4_placeholder_filled):
                        set_labeled_entry(para, "Signed", "", " on behalf of the managers of the responsible hospital")
                        print(f"DEBUG: Filled Signed Part 4 managers at para {i}")
                    # Part 2: "Approved mental health professional"
                    elif "approved mental health" in text_lower:
                        set_labeled_entry(para, "Signed", "", " Approved mental health professional")
                        print(f"DEBUG: Filled Signed AMHP at para {i}")
                    # Part 1 and Part 3: "Responsible clinician"
                    elif "responsible clinician" in text_lower:
                        set_labeled_entry(para, "Signed", "", " Responsible clinician")
                        print(f"DEBUG: Filled Signed RC at para {i}")
                    else:
                        # Default - Responsible clinician
                        set_labeled_entry(para, "Signed", "", " Responsible clinician")
                        print(f"DEBUG: Filled Signed (default) at para {i}")
                    continue

                # Date lines
                if text_stripped.startswith("Date") and "Time" not in text:
                    if i < 35:
                        set_labeled_entry(para, "Date", rc_sig_date)
                    elif i < 50:
                        set_labeled_entry(para, "Date", amhp_sig_date)
                    else:
                        set_labeled_entry(para, "Date", final_sig_date)
                    print(f"DEBUG: Filled Date at para {i}")
                    continue

                # PRINT NAME line (Part 4)
                if "print name" in text_lower:
                    # Handle PRINT NAME [ ] Date [ ] line
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    if para.runs:
                        para.runs[0].text = 'PRINT NAME'
                        pn = para.runs[0]
                    else:
                        pn = para.add_run('PRINT NAME')
                    pn.font.name = 'Arial'
                    pn.font.size = Pt(12)
                    ob1 = para.add_run('[')
                    ob1.font.bold = True
                    ob1.font.color.rgb = BRACKET_COLOR
                    rPr_ob1 = ob1._element.get_or_add_rPr()
                    shd_ob1 = OxmlElement('w:shd')
                    shd_ob1.set(qn('w:val'), 'clear')
                    shd_ob1.set(qn('w:color'), 'auto')
                    shd_ob1.set(qn('w:fill'), CREAM_FILL)
                    rPr_ob1.append(shd_ob1)
                    nm = para.add_run('                              ')
                    rPr_nm = nm._element.get_or_add_rPr()
                    shd_nm = OxmlElement('w:shd')
                    shd_nm.set(qn('w:val'), 'clear')
                    shd_nm.set(qn('w:color'), 'auto')
                    shd_nm.set(qn('w:fill'), CREAM_FILL)
                    rPr_nm.append(shd_nm)
                    cb1 = para.add_run(']')
                    cb1.font.bold = True
                    cb1.font.color.rgb = BRACKET_COLOR
                    rPr_cb1 = cb1._element.get_or_add_rPr()
                    shd_cb1 = OxmlElement('w:shd')
                    shd_cb1.set(qn('w:val'), 'clear')
                    shd_cb1.set(qn('w:color'), 'auto')
                    shd_cb1.set(qn('w:fill'), CREAM_FILL)
                    rPr_cb1.append(shd_cb1)
                    dl = para.add_run(' Date')
                    dl.font.name = 'Arial'
                    dl.font.size = Pt(12)
                    ob2 = para.add_run('[')
                    ob2.font.bold = True
                    ob2.font.color.rgb = BRACKET_COLOR
                    rPr_ob2 = ob2._element.get_or_add_rPr()
                    shd_ob2 = OxmlElement('w:shd')
                    shd_ob2.set(qn('w:val'), 'clear')
                    shd_ob2.set(qn('w:color'), 'auto')
                    shd_ob2.set(qn('w:fill'), CREAM_FILL)
                    rPr_ob2.append(shd_ob2)
                    dt = para.add_run('                              ')
                    rPr_dt = dt._element.get_or_add_rPr()
                    shd_dt = OxmlElement('w:shd')
                    shd_dt.set(qn('w:val'), 'clear')
                    shd_dt.set(qn('w:color'), 'auto')
                    shd_dt.set(qn('w:fill'), CREAM_FILL)
                    rPr_dt.append(shd_dt)
                    cb2 = para.add_run(']')
                    cb2.font.bold = True
                    cb2.font.color.rgb = BRACKET_COLOR
                    rPr_cb2 = cb2._element.get_or_add_rPr()
                    shd_cb2 = OxmlElement('w:shd')
                    shd_cb2.set(qn('w:val'), 'clear')
                    shd_cb2.set(qn('w:color'), 'auto')
                    shd_cb2.set(qn('w:fill'), CREAM_FILL)
                    rPr_cb2.append(shd_cb2)
                    print(f"DEBUG: Filled PRINT NAME/Date at para {i}")
                    continue

                # Furnishing method options (Part 3) - same bracket structure as (i), (ii), (iii)
                # First option gets opening bracket, middle option no brackets, last option gets closing bracket
                if "today consigning it to the hospital" in text_lower:
                    strike_this = any_furnish_selected and not self.furnish_internal.isChecked()
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Opening bracket
                    if para.runs:
                        para.runs[0].text = '['
                        ob = para.runs[0]
                    else:
                        ob = para.add_run('[')
                    ob.font.name = 'Arial'
                    ob.font.size = Pt(12)
                    ob.font.bold = True
                    ob.font.color.rgb = BRACKET_COLOR
                    rPr_ob = ob._element.get_or_add_rPr()
                    shd_ob = OxmlElement('w:shd')
                    shd_ob.set(qn('w:val'), 'clear')
                    shd_ob.set(qn('w:color'), 'auto')
                    shd_ob.set(qn('w:fill'), CREAM_FILL)
                    rPr_ob.append(shd_ob)
                    if strike_this:
                        ob.font.strike = True
                    # Content
                    ct = para.add_run("today consigning it to the hospital managers' internal mail system.")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    if strike_this:
                        ct.font.strike = True
                    print(f"DEBUG: Formatted Part 3 furnish internal at para {i}")
                    continue

                if "today sending it to the hospital" in text_lower:
                    strike_this = any_furnish_selected and not self.furnish_electronic.isChecked()
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Content only - no brackets (middle option)
                    if para.runs:
                        para.runs[0].text = "today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication."
                        ct = para.runs[0]
                    else:
                        ct = para.add_run("today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    if strike_this:
                        ct.font.strike = True
                    print(f"DEBUG: Formatted Part 3 furnish electronic at para {i}")
                    continue

                if "sending or delivering it without" in text_lower:
                    strike_this = any_furnish_selected and not self.furnish_other.isChecked()
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Content
                    if para.runs:
                        para.runs[0].text = "sending or delivering it without using the hospital managers' internal mail system."
                        ct = para.runs[0]
                    else:
                        ct = para.add_run("sending or delivering it without using the hospital managers' internal mail system.")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    if strike_this:
                        ct.font.strike = True
                    # Closing bracket
                    cb = para.add_run(']')
                    cb.font.name = 'Arial'
                    cb.font.size = Pt(12)
                    cb.font.bold = True
                    cb.font.color.rgb = BRACKET_COLOR
                    rPr_cb = cb._element.get_or_add_rPr()
                    shd_cb = OxmlElement('w:shd')
                    shd_cb.set(qn('w:val'), 'clear')
                    shd_cb.set(qn('w:color'), 'auto')
                    shd_cb.set(qn('w:fill'), CREAM_FILL)
                    rPr_cb.append(shd_cb)
                    if strike_this:
                        cb.font.strike = True
                    print(f"DEBUG: Formatted Part 3 furnish other at para {i}")
                    continue

                # Part 4 furnishing options - NO strikethrough, just golden brackets and highlight
                if "furnished to the hospital managers through" in text_lower and "internal mail" in text_lower:
                    # Add opening bracket and cream highlight (no strikethrough in Part 4)
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Opening bracket
                    if para.runs:
                        para.runs[0].text = '['
                        ob = para.runs[0]
                    else:
                        ob = para.add_run('[')
                    ob.font.name = 'Arial'
                    ob.font.size = Pt(12)
                    ob.font.bold = True
                    ob.font.color.rgb = BRACKET_COLOR
                    rPr_ob = ob._element.get_or_add_rPr()
                    shd_ob = OxmlElement('w:shd')
                    shd_ob.set(qn('w:val'), 'clear')
                    shd_ob.set(qn('w:color'), 'auto')
                    shd_ob.set(qn('w:fill'), CREAM_FILL)
                    rPr_ob.append(shd_ob)
                    # Content
                    ct = para.add_run("furnished to the hospital managers through their internal mail system.")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    print(f"DEBUG: Formatted Part 4 furnish internal at para {i}")
                    continue

                if "furnished to the hospital managers, or a person" in text_lower:
                    # Cream highlight only (no strikethrough in Part 4)
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Content with highlight
                    if para.runs:
                        para.runs[0].text = "furnished to the hospital managers, or a person authorised by them to receive it, by means of electronic communication."
                        ct = para.runs[0]
                    else:
                        ct = para.add_run("furnished to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    print(f"DEBUG: Formatted Part 4 furnish electronic at para {i}")
                    continue

                # "received by me on behalf of the hospital managers on [date]."
                # This starts a bracketed section that ends with a placeholder below
                if "received by me on behalf of the hospital managers" in text_lower:
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)
                    # Content with highlight
                    if para.runs:
                        para.runs[0].text = "received by me on behalf of the hospital managers on [date]."
                        ct = para.runs[0]
                    else:
                        ct = para.add_run("received by me on behalf of the hospital managers on [date].")
                    ct.font.name = 'Arial'
                    ct.font.size = Pt(12)
                    rPr_ct = ct._element.get_or_add_rPr()
                    shd_ct = OxmlElement('w:shd')
                    shd_ct.set(qn('w:val'), 'clear')
                    shd_ct.set(qn('w:color'), 'auto')
                    shd_ct.set(qn('w:fill'), CREAM_FILL)
                    rPr_ct.append(shd_ct)
                    part4_received_by_me_found = True
                    print(f"DEBUG: Formatted 'received by me' at para {i}, set flag")
                    continue

                # Part 4 closing bracket placeholder - after "received by me" line
                # This should have cream highlight content with just closing bracket at end
                if part4_received_by_me_found and not part4_placeholder_filled:
                    is_part4_placeholder = len(text) > 10 and (not text_stripped or text.isspace() or text_stripped.replace(' ', '') == '')
                    if is_part4_placeholder:
                        # Part 4 placeholder - content with closing bracket only
                        for run in para.runs:
                            run.text = ""
                        while len(para.runs) > 1:
                            para._element.remove(para.runs[-1]._element)
                        pPr = para._element.get_or_add_pPr()
                        for old_shd in pPr.findall(qn('w:shd')):
                            pPr.remove(old_shd)
                        # Content (placeholder space)
                        if para.runs:
                            para.runs[0].text = '                                                       '
                            ct = para.runs[0]
                        else:
                            ct = para.add_run('                                                       ')
                        ct.font.name = 'Arial'
                        ct.font.size = Pt(12)
                        rPr_ct = ct._element.get_or_add_rPr()
                        shd_ct = OxmlElement('w:shd')
                        shd_ct.set(qn('w:val'), 'clear')
                        shd_ct.set(qn('w:color'), 'auto')
                        shd_ct.set(qn('w:fill'), CREAM_FILL)
                        rPr_ct.append(shd_ct)
                        # Closing bracket
                        cb = para.add_run(']')
                        cb.font.name = 'Arial'
                        cb.font.size = Pt(12)
                        cb.font.bold = True
                        cb.font.color.rgb = BRACKET_COLOR
                        rPr_cb = cb._element.get_or_add_rPr()
                        shd_cb = OxmlElement('w:shd')
                        shd_cb.set(qn('w:val'), 'clear')
                        shd_cb.set(qn('w:color'), 'auto')
                        shd_cb.set(qn('w:fill'), CREAM_FILL)
                        rPr_cb.append(shd_cb)
                        part4_placeholder_filled = True
                        print(f"DEBUG: Formatted Part 4 closing bracket placeholder at para {i}")
                        continue

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO7 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        return {
            "age": self.age_spin.value(),
            "gender": "male" if self.gender_male.isChecked() else "female" if self.gender_female.isChecked() else "other" if self.gender_other.isChecked() else "",
            "ethnicity": self.ethnicity_combo.currentText(),
            "hospital": self.hospital.text(),
            "rc_name": self.rc_name.text(),
            "rc_address": self.rc_address.text(),
            "rc_email": self.rc_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "cto_date": self.cto_date.date().toString("yyyy-MM-dd"),
            "exam_date": self.exam_date.date().toString("yyyy-MM-dd"),
            "dx_primary": self.dx_primary.currentText(),
            "dx_secondary": self.dx_secondary.currentText(),
            "nature": self.nature_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "health": self.health_cb.isChecked(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "others": self.others_cb.isChecked(),
            "cto_effective": self.cto_effective_cb.isChecked(),
            "conditions_met": self.conditions_met_cb.isChecked(),
            "needs_supervision": self.needs_supervision_cb.isChecked(),
            "grounds": self.grounds_preview.text(),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.text(),
            "amhp_email": self.amhp_email.text(),
            "amhp_authority": self.amhp_authority.text(),
            "amhp_approved_by": self.amhp_approved_by.text(),
            "amhp_sig_date": self.amhp_sig_date.date().toString("yyyy-MM-dd"),
            "consulted_name": self.consulted_name.text(),
            "consulted_profession": self.consulted_profession.currentText(),
            "final_sig_date": self.final_sig_date.date().toString("yyyy-MM-dd"),
            "furnish_internal": self.furnish_internal.isChecked(),
            "furnish_electronic": self.furnish_electronic.isChecked(),
            "furnish_other": self.furnish_other.isChecked(),
            "cards": {key: card.get_content() for key, card in self.cards.items()},
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.age_spin.setValue(state.get("age", 0))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        idx = self.ethnicity_combo.findText(state.get("ethnicity", "Ethnicity"))
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.hospital.setText(state.get("hospital", ""))
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_address.setText(state.get("rc_address", ""))
        self.rc_email.setText(state.get("rc_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        if state.get("cto_date"):
            self.cto_date.setDate(QDate.fromString(state["cto_date"], "yyyy-MM-dd"))
        if state.get("exam_date"):
            self.exam_date.setDate(QDate.fromString(state["exam_date"], "yyyy-MM-dd"))
        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)
        dx_secondary = state.get("dx_secondary", "")
        idx = self.dx_secondary.findText(dx_secondary)
        if idx >= 0:
            self.dx_secondary.setCurrentIndex(idx)
        else:
            self.dx_secondary.setCurrentText(dx_secondary)
        self.nature_cb.setChecked(state.get("nature", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.health_cb.setChecked(state.get("health", False))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.cto_effective_cb.setChecked(state.get("cto_effective", False))
        self.conditions_met_cb.setChecked(state.get("conditions_met", False))
        self.needs_supervision_cb.setChecked(state.get("needs_supervision", False))
        self.grounds_preview.setText(state.get("grounds", ""))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setText(state.get("amhp_address", ""))
        self.amhp_email.setText(state.get("amhp_email", ""))
        self.amhp_authority.setText(state.get("amhp_authority", ""))
        self.amhp_approved_by.setText(state.get("amhp_approved_by", ""))
        if state.get("amhp_sig_date"):
            self.amhp_sig_date.setDate(QDate.fromString(state["amhp_sig_date"], "yyyy-MM-dd"))
        self.consulted_name.setText(state.get("consulted_name", ""))
        prof_idx = self.consulted_profession.findText(state.get("consulted_profession", ""))
        if prof_idx >= 0:
            self.consulted_profession.setCurrentIndex(prof_idx)
        if state.get("final_sig_date"):
            self.final_sig_date.setDate(QDate.fromString(state["final_sig_date"], "yyyy-MM-dd"))
        # Restore furnish method
        if state.get("furnish_internal", True):
            self.furnish_internal.setChecked(True)
        elif state.get("furnish_electronic", False):
            self.furnish_electronic.setChecked(True)
        elif state.get("furnish_other", False):
            self.furnish_other.setChecked(True)
        # Restore card contents
        cards_state = state.get("cards", {})
        for key, content in cards_state.items():
            if key in self.cards:
                self.cards[key].set_content_text(content)

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[CTO7Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[CTO7Form] Set gender: {gender}")
        # Set age if not already set (from age field or calculate from DOB)
        if hasattr(self, 'age_spin') and self.age_spin.value() == 0:
            age = patient_info.get("age")
            if not age and patient_info.get("dob"):
                from datetime import datetime
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if age and 0 < age < 120:
                self.age_spin.setValue(age)
                print(f"[CTO7Form] Set age: {age}")
        # Set ethnicity if not already selected
        if patient_info.get("ethnicity") and hasattr(self, 'ethnicity_combo'):
            current = self.ethnicity_combo.currentText()
            if current in ("Ethnicity", "Not specified", "Select..."):
                idx = self.ethnicity_combo.findText(patient_info["ethnicity"], Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    self.ethnicity_combo.setCurrentIndex(idx)
                    print(f"[CTO7Form] Set ethnicity: {patient_info['ethnicity']}")
