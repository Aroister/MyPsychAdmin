# ================================================================
#  A2 FORM PAGE — Section 2 Application by AMHP
#  Mental Health Act 1983 - Form A2 Regulation 4(1)(a)(ii)
#  CARD/POPUP LAYOUT VERSION
# ================================================================

from __future__ import annotations
import re
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QRadioButton, QButtonGroup, QCheckBox, QPushButton,
    QFileDialog, QMessageBox, QGroupBox, QToolButton,
    QStackedWidget, QSplitter
)
from background_history_popup import ResizableSection
from shared_widgets import create_zoom_row
from utils.resource_path import resource_path
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# CARD WIDGET FOR A2 FORM
# ================================================================
class A2CardWidget(QFrame):
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.title = title
        self.key = key

        self.setObjectName("a2Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#a2Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#a2Card:hover {
                border-color: #2563eb;
                background: #eff6ff;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QLabel#cardTitle {
                font-size: 20px;
                font-weight: 600;
                color: #2563eb;
                padding-bottom: 4px;
            }
            QFrame#divider {
                background: #e5e7eb;
                height: 1px;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 12)
        layout.setSpacing(0)

        # Fixed header section with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setObjectName("cardTitle")
        self.title_label.setFixedHeight(24)
        self.title_label.setCursor(Qt.CursorShape.PointingHandCursor)
        self.title_label.mousePressEvent = lambda e: self.clicked.emit(self.key)
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        divider = QFrame()
        divider.setObjectName("divider")
        divider.setFixedHeight(1)
        layout.addWidget(divider)

        # Scrollable content area
        self.content_scroll = QScrollArea()
        self.content_scroll.setWidgetResizable(True)
        self.content_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        self.content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.content_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        content_widget = QWidget()
        content_widget.setStyleSheet("background: transparent;")
        content_layout = QVBoxLayout(content_widget)
        content_layout.setContentsMargins(0, 8, 0, 4)
        content_layout.setSpacing(0)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 20px;
                color: #374151;
                padding: 4px;
                background: transparent;
                border: none;
            }
        """)
        self.content.setFrameShape(QFrame.Shape.NoFrame)
        content_layout.addWidget(self.content)
        content_layout.addStretch()

        self.content_scroll.setWidget(content_widget)
        layout.addWidget(self.content_scroll, 1)

        # Add zoom controls to header row (after content is created)
        zoom_row = create_zoom_row(self.content, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def set_content(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)


# ================================================================
# MAIN A2 FORM PAGE
# ================================================================
class A2FormPage(QWidget):
    """Page for completing MHA Form A2 - Section 2 AMHP Application."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}
        self._active_section = None
        self._setup_ui()
        self._prefill_amhp_details()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {
            "full_name": details[1] or "",
            "email": details[7] or "",
        }

    def _prefill_amhp_details(self):
        if self._my_details.get("full_name"):
            self.amhp_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.amhp_email.setText(self._my_details["email"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("QFrame { background: #2563eb; border-bottom: 1px solid #1d4ed8; }")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Form A2 — Section 2 Application by AMHP")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover {
                background: #7f1d1d;
            }
            QPushButton:pressed {
                background: #450a0a;
            }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Split layout: cards on left, popups on right
        split = QSplitter(Qt.Orientation.Horizontal)
        split.setHandleWidth(6)
        split.setStyleSheet("""
            QSplitter::handle { background: #d1d5db; }
            QSplitter::handle:hover { background: #6BAF8D; }
        """)
        main_layout.addWidget(split, 1)

        # Sections for cards
        self.sections = [
            ("Hospital", "hospital"),
            ("AMHP Details", "amhp"),
            ("Patient Details", "patient"),
            ("Local Authority", "authority"),
            ("Nearest Relative", "nr"),
            ("Patient Interview", "interview"),
            ("Medical Recommendations", "medical"),
            ("Signature", "signature"),
        ]

        # ---------------- LEFT: CARDS ----------------
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("QScrollArea { background: #f9fafb; border: none; }")
        split.addWidget(self.cards_holder)

        cards_container = QWidget()
        cards_container.setStyleSheet("background: #f9fafb;")
        cards_layout = QVBoxLayout(cards_container)
        cards_layout.setContentsMargins(30, 30, 30, 30)
        cards_layout.setSpacing(20)

        self.card_sections = {}
        for title, key in self.sections:
            section = ResizableSection()
            section.set_content_height(180)
            section._min_height = 120
            section._max_height = 350

            card = A2CardWidget(title, key)
            self._hook_editor_focus(card.content)
            card.clicked.connect(self._activate_section)
            self.cards[key] = card
            section.set_content(card)
            self.card_sections[key] = section
            cards_layout.addWidget(section)

        cards_layout.addStretch()
        self.cards_holder.setWidget(cards_container)

        # ---------------- RIGHT: POPUP PANEL ----------------
        self.popup_panel = QFrame()
        self.popup_panel.setMinimumWidth(400)
        self.popup_panel.setMaximumWidth(650)
        self.popup_panel.setStyleSheet("QFrame { background: rgba(245,245,245,0.95); border-left: 1px solid rgba(0,0,0,0.08); }")
        split.addWidget(self.popup_panel)
        split.setSizes([500, 480])

        panel_layout = QVBoxLayout(self.popup_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setStyleSheet("""
            font-size: 22px; font-weight: 700; color: #2563eb;
            background: rgba(219,234,254,0.85); padding: 8px 12px; border-radius: 8px;
        """)
        panel_layout.addWidget(self.panel_title)

        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        self._build_popup_panels()
        self._activate_section("hospital")

        # Initialize cards with default date values
        self._update_interview_card()
        self._update_signature_card()

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _activate_section(self, key: str):
        self._active_section = key
        index = [s[1] for s in self.sections].index(key)
        self.popup_stack.setCurrentIndex(index)
        title = [s[0] for s in self.sections if s[1] == key][0]
        self.panel_title.setText(title)

    def _build_popup_panels(self):
        self._build_hospital_popup()
        self._build_amhp_popup()
        self._build_patient_popup()
        self._build_authority_popup()
        self._build_nr_popup()
        self._build_interview_popup()
        self._build_medical_popup()
        self._build_signature_popup()

    def _create_popup_scroll(self) -> tuple:
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setMaximumWidth(440)
        layout = QVBoxLayout(container)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(12)

        scroll.setWidget(container)
        return scroll, layout

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 17px; }
            QLineEdit:focus { border-color: #2563eb; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 80) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMaximumHeight(height)
        edit.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 8px; font-size: 17px; }
            QTextEdit:focus { border-color: #2563eb; }
        """)
        return edit

    def _create_section_label(self, text: str) -> QLabel:
        lbl = QLabel(text)
        lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        return lbl

    # ----------------------------------------------------------------
    # POPUP BUILDERS
    # ----------------------------------------------------------------
    def _build_hospital_popup(self):
        scroll, layout = self._create_popup_scroll()
        layout.addWidget(self._create_section_label("Hospital Name:"))
        self.hospital_name = self._create_line_edit("Enter hospital name")
        layout.addWidget(self.hospital_name)
        layout.addWidget(self._create_section_label("Hospital Address:"))
        self.hospital_address = self._create_text_edit("Enter hospital address")
        layout.addWidget(self.hospital_address)

        # Auto-sync to card
        self.hospital_name.textChanged.connect(self._update_hospital_card)
        self.hospital_address.textChanged.connect(self._update_hospital_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_hospital_card(self):
        parts = []
        if self.hospital_name.text().strip():
            parts.append(self.hospital_name.text().strip())
        if self.hospital_address.toPlainText().strip():
            parts.append(self.hospital_address.toPlainText().strip()[:50] + "...")
        self.cards["hospital"].set_content("\n".join(parts) if parts else "Click to enter details")

    def _build_amhp_popup(self):
        scroll, layout = self._create_popup_scroll()
        layout.addWidget(self._create_section_label("AMHP Full Name:"))
        self.amhp_name = self._create_line_edit("Enter your full name")
        layout.addWidget(self.amhp_name)
        layout.addWidget(self._create_section_label("Address:"))
        self.amhp_address = self._create_text_edit("Enter your address")
        layout.addWidget(self.amhp_address)
        layout.addWidget(self._create_section_label("Email:"))
        self.amhp_email = self._create_line_edit("Enter email address")
        layout.addWidget(self.amhp_email)

        # Auto-sync to card
        self.amhp_name.textChanged.connect(self._update_amhp_card)
        self.amhp_address.textChanged.connect(self._update_amhp_card)
        self.amhp_email.textChanged.connect(self._update_amhp_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_amhp_card(self):
        parts = []
        if self.amhp_name.text().strip():
            parts.append(self.amhp_name.text().strip())
        if self.amhp_address.toPlainText().strip():
            parts.append(self.amhp_address.toPlainText().strip()[:50] + "...")
        if self.amhp_email.text().strip():
            parts.append(self.amhp_email.text().strip())
        self.cards["amhp"].set_content("\n".join(parts) if parts else "Click to enter details")

    def _build_patient_popup(self):
        scroll, layout = self._create_popup_scroll()
        layout.addWidget(self._create_section_label("Patient Full Name:"))
        self.patient_name = self._create_line_edit("Enter patient's full name")
        layout.addWidget(self.patient_name)
        layout.addWidget(self._create_section_label("Patient Address:"))
        self.patient_address = self._create_text_edit("Enter patient's address")
        layout.addWidget(self.patient_address)

        # Auto-sync to card
        self.patient_name.textChanged.connect(self._update_patient_card)
        self.patient_address.textChanged.connect(self._update_patient_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_patient_card(self):
        parts = []
        if self.patient_name.text().strip():
            parts.append(self.patient_name.text().strip())
        if self.patient_address.toPlainText().strip():
            parts.append(self.patient_address.toPlainText().strip()[:50] + "...")
        self.cards["patient"].set_content("\n".join(parts) if parts else "Click to enter details")

    def _build_authority_popup(self):
        scroll, layout = self._create_popup_scroll()
        layout.addWidget(self._create_section_label("Local Social Services Authority:"))
        self.local_authority = self._create_line_edit("Enter local authority")
        layout.addWidget(self.local_authority)

        # Approved by group
        group = QFrame()
        group.setStyleSheet("QFrame { background: #f0f9ff; border: none; border-radius: 8px; }")
        group_layout = QVBoxLayout(group)
        group_layout.setContentsMargins(12, 10, 12, 10)

        group_header = QLabel("Approved by:")
        group_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #0369a1;")
        group_layout.addWidget(group_header)

        self.approved_by_same = QRadioButton("That authority (same as above)")
        self.approved_by_different = QRadioButton("Different authority:")
        self.approved_by_same.setChecked(True)
        self.approved_btn_group = QButtonGroup()
        self.approved_btn_group.addButton(self.approved_by_same)
        self.approved_btn_group.addButton(self.approved_by_different)

        group_layout.addWidget(self.approved_by_same)
        group_layout.addWidget(self.approved_by_different)

        self.approved_by_authority = self._create_line_edit("Enter approving authority if different")
        self.approved_by_authority.setEnabled(False)
        self.approved_by_different.toggled.connect(self.approved_by_authority.setEnabled)
        group_layout.addWidget(self.approved_by_authority)

        layout.addWidget(group)

        # Auto-sync to card
        self.local_authority.textChanged.connect(self._update_authority_card)
        self.approved_by_same.toggled.connect(self._update_authority_card)
        self.approved_by_authority.textChanged.connect(self._update_authority_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_authority_card(self):
        parts = []
        if self.local_authority.text().strip():
            parts.append(self.local_authority.text().strip())
        if self.approved_by_different.isChecked() and self.approved_by_authority.text().strip():
            parts.append(self.approved_by_authority.text().strip())
        self.cards["authority"].set_content("\n".join(parts) if parts else "Click to enter details")

    def _build_nr_popup(self):
        scroll, layout = self._create_popup_scroll()

        # Known/Unknown
        known_frame = QFrame()
        known_frame.setStyleSheet("QFrame { background: #fef3c7; border: none; border-radius: 8px; }")
        known_layout = QVBoxLayout(known_frame)
        known_layout.setContentsMargins(12, 10, 12, 10)

        known_header = QLabel("Do you know who the nearest relative is?")
        known_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #92400e;")
        known_layout.addWidget(known_header)

        self.nr_known = QRadioButton("Yes - I know who the nearest relative is")
        self.nr_unknown = QRadioButton("No - I don't know / patient has no NR")
        self.nr_known.setChecked(True)
        self.known_btn_group = QButtonGroup()
        self.known_btn_group.addButton(self.nr_known)
        self.known_btn_group.addButton(self.nr_unknown)

        known_layout.addWidget(self.nr_known)
        known_layout.addWidget(self.nr_unknown)
        layout.addWidget(known_frame)

        # NR Details (when known)
        self.nr_details_widget = QWidget()
        nr_details_layout = QVBoxLayout(self.nr_details_widget)
        nr_details_layout.setContentsMargins(0, 8, 0, 0)
        nr_details_layout.setSpacing(8)

        # Option (a) or (b)
        option_frame = QFrame()
        option_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 8px; }")
        option_layout = QVBoxLayout(option_frame)
        option_layout.setContentsMargins(12, 10, 12, 10)

        self.nr_option_a = QRadioButton("(a) This person IS the patient's nearest relative")
        self.nr_option_b = QRadioButton("(b) This person has been AUTHORISED to act as NR")
        self.nr_option_a.setChecked(True)
        self.option_btn_group = QButtonGroup()
        self.option_btn_group.addButton(self.nr_option_a)
        self.option_btn_group.addButton(self.nr_option_b)

        option_layout.addWidget(self.nr_option_a)
        option_layout.addWidget(self.nr_option_b)
        nr_details_layout.addWidget(option_frame)

        nr_details_layout.addWidget(self._create_section_label("NR Name:"))
        self.nr_name = self._create_line_edit("Enter nearest relative's full name")
        nr_details_layout.addWidget(self.nr_name)

        nr_details_layout.addWidget(self._create_section_label("NR Address:"))
        self.nr_address = self._create_text_edit("Enter nearest relative's address")
        nr_details_layout.addWidget(self.nr_address)

        # Informed
        informed_frame = QFrame()
        informed_frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 8px; }")
        informed_layout = QVBoxLayout(informed_frame)
        informed_layout.setContentsMargins(12, 10, 12, 10)

        informed_header = QLabel("Have you informed this person?")
        informed_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1e40af;")
        informed_layout.addWidget(informed_header)

        self.nr_informed_yes = QRadioButton("Yes - I have informed them")
        self.nr_informed_no = QRadioButton("No - I have not yet informed them")
        self.nr_informed_yes.setChecked(True)
        self.informed_btn_group = QButtonGroup()
        self.informed_btn_group.addButton(self.nr_informed_yes)
        self.informed_btn_group.addButton(self.nr_informed_no)

        informed_layout.addWidget(self.nr_informed_yes)
        informed_layout.addWidget(self.nr_informed_no)
        nr_details_layout.addWidget(informed_frame)

        layout.addWidget(self.nr_details_widget)

        # Unknown NR section
        self.nr_unknown_widget = QWidget()
        nr_unknown_layout = QVBoxLayout(self.nr_unknown_widget)
        nr_unknown_layout.setContentsMargins(0, 8, 0, 0)

        unknown_frame = QFrame()
        unknown_frame.setStyleSheet("QFrame { background: #fef2f2; border: none; border-radius: 8px; }")
        unknown_opt_layout = QVBoxLayout(unknown_frame)
        unknown_opt_layout.setContentsMargins(12, 10, 12, 10)

        unknown_header = QLabel("Select option:")
        unknown_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #991b1b;")
        unknown_opt_layout.addWidget(unknown_header)

        self.nr_unable = QRadioButton("(a) Unable to ascertain who is the NR")
        self.nr_none = QRadioButton("(b) Patient has no NR within meaning of the Act")
        self.nr_unable.setChecked(True)
        self.unknown_btn_group = QButtonGroup()
        self.unknown_btn_group.addButton(self.nr_unable)
        self.unknown_btn_group.addButton(self.nr_none)

        unknown_opt_layout.addWidget(self.nr_unable)
        unknown_opt_layout.addWidget(self.nr_none)
        nr_unknown_layout.addWidget(unknown_frame)

        layout.addWidget(self.nr_unknown_widget)
        self.nr_unknown_widget.hide()

        self.nr_known.toggled.connect(self._toggle_nr_sections)

        # Auto-sync to card
        self.nr_known.toggled.connect(self._update_nr_card)
        self.nr_option_a.toggled.connect(self._update_nr_card)
        self.nr_name.textChanged.connect(self._update_nr_card)
        self.nr_address.textChanged.connect(self._update_nr_card)
        self.nr_informed_yes.toggled.connect(self._update_nr_card)
        self.nr_unable.toggled.connect(self._update_nr_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _toggle_nr_sections(self, known: bool):
        self.nr_details_widget.setVisible(known)
        self.nr_unknown_widget.setVisible(not known)

    def _update_nr_card(self):
        parts = []
        if self.nr_known.isChecked():
            if self.nr_name.text().strip():
                parts.append(self.nr_name.text().strip())
            if self.nr_address.toPlainText().strip():
                parts.append(self.nr_address.toPlainText().strip()[:50] + "...")
        self.cards["nr"].set_content("\n".join(parts) if parts else "Click to enter details")

    def _build_interview_popup(self):
        scroll, layout = self._create_popup_scroll()
        layout.addWidget(self._create_section_label("Date patient last seen:"))

        self.last_seen_date = QDateEdit()
        self.last_seen_date.setCalendarPopup(True)
        self.last_seen_date.setDate(QDate.currentDate())
        self.last_seen_date.setStyleSheet("""
            QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 17px; }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        layout.addWidget(self.last_seen_date)

        info = QLabel("This must be within 14 days ending on the day this application is completed.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 20px; color: #6b7280; padding: 8px; background: #f3f4f6; border-radius: 6px;")
        layout.addWidget(info)

        # Auto-sync to card
        self.last_seen_date.dateChanged.connect(self._update_interview_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_interview_card(self):
        date_str = self.last_seen_date.date().toString("dd MMM yyyy")
        self.cards["interview"].set_content(date_str)

    def _build_medical_popup(self):
        scroll, layout = self._create_popup_scroll()

        info = QLabel("This application is founded on two medical recommendations in the prescribed form.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 17px; color: #374151; padding: 12px; background: #dbeafe; border-radius: 6px;")
        layout.addWidget(info)

        reason_lbl = QLabel("If neither medical practitioner had previous acquaintance with the patient, explain why:")
        reason_lbl.setWordWrap(True)
        reason_lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151; margin-top: 8px;")
        layout.addWidget(reason_lbl)

        self.no_acquaintance_reason = QTextEdit()
        self.no_acquaintance_reason.setPlaceholderText("Enter explanation (leave blank if not applicable)")
        self.no_acquaintance_reason.setMinimumHeight(100)
        self.no_acquaintance_reason.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 8px; font-size: 17px; }
            QTextEdit:focus { border-color: #2563eb; }
        """)
        layout.addWidget(self.no_acquaintance_reason)

        # Auto-sync to card
        self.no_acquaintance_reason.textChanged.connect(self._update_medical_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_medical_card(self):
        if self.no_acquaintance_reason.toPlainText().strip():
            text = self.no_acquaintance_reason.toPlainText().strip()[:80] + "..."
            self.cards["medical"].set_content(text)
        else:
            self.cards["medical"].set_content("Click to enter details")

    def _build_signature_popup(self):
        scroll, layout = self._create_popup_scroll()
        layout.addWidget(self._create_section_label("Signature Date:"))

        self.signature_date = QDateEdit()
        self.signature_date.setCalendarPopup(True)
        self.signature_date.setDate(QDate.currentDate())
        self.signature_date.setStyleSheet("""
            QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 17px; }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        layout.addWidget(self.signature_date)

        info = QLabel("The form will be signed manually after printing.")
        info.setStyleSheet("font-size: 20px; color: #6b7280; padding: 8px; background: #f3f4f6; border-radius: 6px;")
        layout.addWidget(info)

        # Auto-sync to card
        self.signature_date.dateChanged.connect(self._update_signature_card)

        layout.addStretch()
        self.popup_stack.addWidget(scroll)

    def _update_signature_card(self):
        date_str = self.signature_date.date().toString("dd MMM yyyy")
        self.cards["signature"].set_content(date_str)

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _go_back(self):
        self.go_back.emit()

    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.local_authority.clear()
            self.approved_by_same.setChecked(True)
            self.approved_by_authority.clear()
            self.nr_known.setChecked(True)
            self.nr_option_a.setChecked(True)
            self.nr_name.clear()
            self.nr_address.clear()
            self.nr_informed_yes.setChecked(True)
            self.nr_unable.setChecked(True)
            self.last_seen_date.setDate(QDate.currentDate())
            self.no_acquaintance_reason.clear()
            self.signature_date.setDate(QDate.currentDate())
            for card in self.cards.values():
                card.set_content("Click to enter details")
            # Restore my details fields
            self._prefill_amhp_details()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Form A2",
            f"Form_A2_{datetime.now().strftime('%Y%m%d')}.docx", "Word Documents (*.docx)")
        if not file_path:
            return
        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_A2_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form A2 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Clean ALL paragraphs - remove permission markers and convert grey to cream
            # This approach matches A4 which works correctly
            for para in doc.paragraphs:
                para_xml = para._element
                # Remove permission markers (they cause grey appearance)
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                # Convert paragraph-level grey shading to cream (keep the shading, just change color)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), 'FFFED5')
                # Convert run-level grey shading to cream
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), 'FFFED5')

            def fill_bracket_content(para, content):
                """Replace placeholder text inside brackets - bold gold brackets, content cream highlighted"""
                full_text = ''.join(run.text for run in para.runs)

                # Find bracket pattern to get text before and after
                bracket_match = re.search(r'\[([^\]]*)\]', full_text)
                if bracket_match:
                    text_before = full_text[:bracket_match.start()]
                    text_after = full_text[bracket_match.end():]
                else:
                    text_before = ""
                    text_after = ""

                # Clear all runs
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)

                # Add text before bracket (if any)
                if text_before:
                    if para.runs:
                        para.runs[0].text = text_before
                        para.runs[0].font.name = 'Arial'
                        para.runs[0].font.size = Pt(12)
                    else:
                        run = para.add_run(text_before)
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                else:
                    if para.runs:
                        para.runs[0].text = ""

                # Add opening bracket with bold gold color
                bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr = bracket_open._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)

                # Add content with cream highlighting
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'FFFED5')
                rPr2.append(shd2)

                # Add closing bracket with bold gold color
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr3 = bracket_close._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)

                # Add text after bracket (if any)
                if text_after:
                    after_run = para.add_run(text_after)
                    after_run.font.name = 'Arial'
                    after_run.font.size = Pt(12)

            def set_para_text(para, new_text):
                """Set paragraph text - if empty, use spaces for visible blank placeholder"""
                if not new_text.strip():
                    new_text = '                                                                   '
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                else:
                    run = para.add_run(new_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            # Gold bracket color - brighter gold #918C0D
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)

            def set_entry_box(para, content=""):
                """Set entry box with bold gold brackets [content] - brackets are gold, content is cream"""
                # If no content, use spaces to create blank entry box
                if not content.strip():
                    content = '                                                                   '

                # Clear existing runs
                for run in para.runs:
                    run.text = ""

                # Remove empty runs except first
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)

                # Add opening bracket with gold color and bold
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                # Add cream shading to bracket too
                rPr = bracket_open._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)

                # Add content with cream highlighting (no color change)
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'FFFED5')
                rPr2.append(shd2)

                # Add closing bracket with gold color and bold
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr3 = bracket_close._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)

            def highlight_yellow(para):
                """Apply cream highlighting to all runs in paragraph (matches A4 approach)"""
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    existing_shd = rPr.find(qn('w:shd'))
                    if existing_shd is not None:
                        # Change existing shading to cream
                        existing_shd.set(qn('w:fill'), 'FFFED5')
                    else:
                        # Add cream shading if none exists
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'FFFED5')
                        rPr.append(shd)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Template structure:
            # Para 2: "To the managers of [name and address of hospital]" - instruction line
            # Para 3: BLANK entry field (30 spaces) - fill with hospital data
            # Para 4: "I [PRINT your full name, address...]" - instruction line
            # Para 5: BLANK entry field - fill with AMHP data
            # Para 6: "apply for the admission of [PRINT full name and address of patient]" - instruction
            # Para 7: BLANK entry field - fill with patient data
            # Para 9: "I am acting on behalf of [PRINT name of local social services authority]" - instruction
            # Para 10: BLANK entry field - fill with authority data

            hospital_text = self.cards["hospital"].get_content().replace("\n", ", ")
            amhp_text = self.cards["amhp"].get_content().replace("\n", ", ")
            patient_text = self.cards["patient"].get_content().replace("\n", ", ")

            # Entry box paragraphs - format as [content] or [blank] with cream highlight
            # Para 3: Hospital entry box
            if hospital_text.strip() and hospital_text != "Click to enter details":
                set_entry_box(paragraphs[3], hospital_text)
            else:
                set_entry_box(paragraphs[3], "")

            # Para 5: AMHP entry box
            if amhp_text.strip() and amhp_text != "Click to enter details":
                set_entry_box(paragraphs[5], amhp_text)
            else:
                set_entry_box(paragraphs[5], "")

            # Para 7: Patient entry box
            if patient_text.strip() and patient_text != "Click to enter details":
                set_entry_box(paragraphs[7], patient_text)
            else:
                set_entry_box(paragraphs[7], "")

            # Para 10: Local authority entry box
            authority_text = self.cards["authority"].get_content().replace("\n", ", ")
            if authority_text.strip() and authority_text != "Click to enter details":
                set_entry_box(paragraphs[10], authority_text)
            else:
                set_entry_box(paragraphs[10], "")

            # Paragraphs 11-14: Authority approval
            # Para 12: "that authority" - wrap in gold brackets with cream highlight
            set_entry_box(paragraphs[12], "that authority")
            # Para 14: Entry box for different authority - ALWAYS show placeholder
            if self.approved_by_same.isChecked():
                # "that authority" is selected - strikethrough the alternative
                strikethrough_para(paragraphs[13])
                set_entry_box(paragraphs[14], "")  # Still show blank placeholder
                strikethrough_para(paragraphs[14])
            else:
                strikethrough_para(paragraphs[12])
                if self.approved_by_authority.text():
                    set_entry_box(paragraphs[14], self.approved_by_authority.text())
                else:
                    set_entry_box(paragraphs[14], "")

            # Nearest relative section
            # Para 17-23: Wrapped in gold brackets (the "known NR" section)
            # Para 25-26: Wrapped in gold brackets (the "unknown NR" section)

            # Apply cream highlighting to text paragraphs that should be highlighted (per template)
            highlight_yellow(paragraphs[17])  # "(a) To the best of my knowledge and belief..."
            highlight_yellow(paragraphs[19])  # "is the patient's nearest relative within the meaning of the Act."
            highlight_yellow(paragraphs[20])  # "(b) I understand that [PRINT full name and address]"
            highlight_yellow(paragraphs[22])  # "has been authorised by a county court..."
            highlight_yellow(paragraphs[23])  # "I have/have not yet*..."

            # Para 25-26: Add gold brackets around the "unknown NR" section
            # Para 25 starts with gold [ and Para 26 ends with gold ]
            # Para 25: "(a) I have been unable to ascertain..."
            p25 = paragraphs[25]
            p25_text = ''.join(run.text for run in p25.runs)
            for run in p25.runs:
                run.text = ""
            while len(p25.runs) > 1:
                p25._element.remove(p25.runs[-1]._element)
            # Add opening gold bracket
            if p25.runs:
                p25.runs[0].text = '['
                p25.runs[0].font.name = 'Arial'
                p25.runs[0].font.size = Pt(12)
                p25.runs[0].font.bold = True
                p25.runs[0].font.color.rgb = BRACKET_COLOR
                rPr = p25.runs[0]._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)
            # Add the text (cream highlighted)
            p25_content = p25.add_run(p25_text)
            p25_content.font.name = 'Arial'
            p25_content.font.size = Pt(12)
            rPr2 = p25_content._element.get_or_add_rPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'), 'clear')
            shd2.set(qn('w:color'), 'auto')
            shd2.set(qn('w:fill'), 'FFFED5')
            rPr2.append(shd2)

            # Para 26: "(b) To the best of my knowledge and belief this patient has no nearest relative..."
            p26 = paragraphs[26]
            p26_text = ''.join(run.text for run in p26.runs)
            for run in p26.runs:
                run.text = ""
            while len(p26.runs) > 1:
                p26._element.remove(p26.runs[-1]._element)
            # Add the text (cream highlighted)
            if p26.runs:
                p26.runs[0].text = p26_text
                p26.runs[0].font.name = 'Arial'
                p26.runs[0].font.size = Pt(12)
                rPr3 = p26.runs[0]._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)
            # Add closing gold bracket
            p26_close = p26.add_run(']')
            p26_close.font.name = 'Arial'
            p26_close.font.size = Pt(12)
            p26_close.font.bold = True
            p26_close.font.color.rgb = BRACKET_COLOR
            rPr4 = p26_close._element.get_or_add_rPr()
            shd4 = OxmlElement('w:shd')
            shd4.set(qn('w:val'), 'clear')
            shd4.set(qn('w:color'), 'auto')
            shd4.set(qn('w:fill'), 'FFFED5')
            rPr4.append(shd4)

            if self.nr_known.isChecked():
                nr_text = self.cards["nr"].get_content().replace("\n", ", ")

                if self.nr_option_a.isChecked():
                    # Option A - is the nearest relative
                    # Para 18: Gold bracketed entry box
                    if nr_text.strip() and nr_text != "Click to enter details":
                        set_entry_box(paragraphs[18], nr_text)
                    else:
                        set_entry_box(paragraphs[18], "")
                    # Strikethrough option B paragraphs - but still show placeholder for para 21
                    strikethrough_para(paragraphs[20])
                    set_entry_box(paragraphs[21], "")
                    strikethrough_para(paragraphs[21])
                    strikethrough_para(paragraphs[22])
                else:
                    # Option B - authorized person
                    # Para 21: Gold bracketed entry box
                    if nr_text.strip() and nr_text != "Click to enter details":
                        set_entry_box(paragraphs[21], nr_text)
                    else:
                        set_entry_box(paragraphs[21], "")
                    # Strikethrough option A paragraphs - but still show placeholder for para 18
                    strikethrough_para(paragraphs[17])
                    set_entry_box(paragraphs[18], "")
                    strikethrough_para(paragraphs[18])
                    strikethrough_para(paragraphs[19])

                # Strikethrough the "unknown NR" section
                strikethrough_para(paragraphs[25])
                strikethrough_para(paragraphs[26])

                para = paragraphs[23]
                for run in para.runs:
                    run.text = ""

                if self.nr_informed_yes.isChecked():
                    run1 = para.add_run("I have")
                    run1.font.name = 'Arial'
                    run1.font.size = Pt(12)
                    run2 = para.add_run("/have not yet*")
                    run2.font.name = 'Arial'
                    run2.font.size = Pt(12)
                    run2.font.strike = True
                    run3 = para.add_run(" informed that person that this application is to be made and of the nearest relative's power to order the discharge of the patient.")
                    run3.font.name = 'Arial'
                    run3.font.size = Pt(12)
                else:
                    run1 = para.add_run("I have/")
                    run1.font.name = 'Arial'
                    run1.font.size = Pt(12)
                    run1.font.strike = True
                    run2 = para.add_run("have not yet")
                    run2.font.name = 'Arial'
                    run2.font.size = Pt(12)
                    run3 = para.add_run("*")
                    run3.font.name = 'Arial'
                    run3.font.size = Pt(12)
                    run3.font.strike = True
                    run4 = para.add_run(" informed that person that this application is to be made and of the nearest relative's power to order the discharge of the patient.")
                    run4.font.name = 'Arial'
                    run4.font.size = Pt(12)
            else:
                # NR unknown - strikethrough the "known NR" section (paras 15-23)
                # But still show placeholders for paras 18 and 21 with gold brackets
                set_entry_box(paragraphs[18], "")
                set_entry_box(paragraphs[21], "")
                for i in range(15, 24):
                    strikethrough_para(paragraphs[i])
                # Select appropriate unknown reason
                if self.nr_unable.isChecked():
                    # Unable to ascertain - strikethrough "no NR" option
                    strikethrough_para(paragraphs[26])
                else:
                    # Patient has no NR - strikethrough "unable to ascertain" option
                    strikethrough_para(paragraphs[25])

            # Paragraph 28: "I last saw the patient on [date]," - keep as-is (label only)
            # Don't modify - the [date] text stays as a label

            # Para 29: Entry box below date line - gold brackets with cream highlight
            last_seen = self.last_seen_date.date().toString("dd MMMM yyyy")
            set_entry_box(paragraphs[29], last_seen)

            # Paragraph 34: No acquaintance reason entry box
            if self.no_acquaintance_reason.toPlainText():
                set_entry_box(paragraphs[34], self.no_acquaintance_reason.toPlainText())
            else:
                set_entry_box(paragraphs[34], "")

            # Para 35: "[If you need to continue on a separate sheet please indicate here [ ] and attach that sheet to this form]"
            # Outer brackets are black, inner [ ] checkbox has bold gold brackets with cream space
            para35 = paragraphs[35]
            # Clear existing runs
            for run in para35.runs:
                run.text = ""
            while len(para35.runs) > 1:
                para35._element.remove(para35.runs[-1]._element)

            # Opening bracket (black, no highlight)
            if para35.runs:
                para35.runs[0].text = '['
                para35.runs[0].font.name = 'Arial'
                para35.runs[0].font.size = Pt(12)

            # Text before checkbox (no highlight)
            text1 = para35.add_run('If you need to continue on a separate sheet please indicate here ')
            text1.font.name = 'Arial'
            text1.font.size = Pt(12)

            # Checkbox opening bracket (bold gold color, cream highlight)
            cb_open = para35.add_run('[')
            cb_open.font.name = 'Arial'
            cb_open.font.size = Pt(12)
            cb_open.font.bold = True
            cb_open.font.color.rgb = BRACKET_COLOR
            rPr1 = cb_open._element.get_or_add_rPr()
            shd1 = OxmlElement('w:shd')
            shd1.set(qn('w:val'), 'clear')
            shd1.set(qn('w:color'), 'auto')
            shd1.set(qn('w:fill'), 'FFFED5')
            rPr1.append(shd1)

            # Checkbox space (cream highlight)
            cb_space = para35.add_run('  ')
            cb_space.font.name = 'Arial'
            cb_space.font.size = Pt(12)
            rPr2 = cb_space._element.get_or_add_rPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'), 'clear')
            shd2.set(qn('w:color'), 'auto')
            shd2.set(qn('w:fill'), 'FFFED5')
            rPr2.append(shd2)

            # Checkbox closing bracket (bold gold color, cream highlight)
            cb_close = para35.add_run(']')
            cb_close.font.name = 'Arial'
            cb_close.font.size = Pt(12)
            cb_close.font.bold = True
            cb_close.font.color.rgb = BRACKET_COLOR
            rPr3 = cb_close._element.get_or_add_rPr()
            shd3 = OxmlElement('w:shd')
            shd3.set(qn('w:val'), 'clear')
            shd3.set(qn('w:color'), 'auto')
            shd3.set(qn('w:fill'), 'FFFED5')
            rPr3.append(shd3)

            # Text after checkbox (no highlight)
            text2 = para35.add_run(' and attach that sheet to this form')
            text2.font.name = 'Arial'
            text2.font.size = Pt(12)

            # Closing bracket (black, no highlight)
            close_bracket = para35.add_run(']')
            close_bracket.font.name = 'Arial'
            close_bracket.font.size = Pt(12)

            # Paragraph 36: Signature line - "Signed [    ] Date [date]"
            sig_date = self.signature_date.date().toString("dd MMMM yyyy")
            sig_para = paragraphs[36]
            # Clear existing runs
            for run in sig_para.runs:
                run.text = ""
            while len(sig_para.runs) > 1:
                sig_para._element.remove(sig_para.runs[-1]._element)

            # "Signed "
            if sig_para.runs:
                sig_para.runs[0].text = "Signed "
                sig_para.runs[0].font.name = 'Arial'
                sig_para.runs[0].font.size = Pt(12)

            # Opening bracket for signature (bold gold color)
            sig_open = sig_para.add_run('[')
            sig_open.font.name = 'Arial'
            sig_open.font.size = Pt(12)
            sig_open.font.bold = True
            sig_open.font.color.rgb = BRACKET_COLOR
            rPr = sig_open._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'FFFED5')
            rPr.append(shd)

            # Signature space (cream highlighted)
            sig_space = sig_para.add_run('                                        ')
            sig_space.font.name = 'Arial'
            sig_space.font.size = Pt(12)
            rPr2 = sig_space._element.get_or_add_rPr()
            shd2 = OxmlElement('w:shd')
            shd2.set(qn('w:val'), 'clear')
            shd2.set(qn('w:color'), 'auto')
            shd2.set(qn('w:fill'), 'FFFED5')
            rPr2.append(shd2)

            # Closing bracket for signature (bold gold color)
            sig_close = sig_para.add_run(']')
            sig_close.font.name = 'Arial'
            sig_close.font.size = Pt(12)
            sig_close.font.bold = True
            sig_close.font.color.rgb = BRACKET_COLOR
            rPr3 = sig_close._element.get_or_add_rPr()
            shd3 = OxmlElement('w:shd')
            shd3.set(qn('w:val'), 'clear')
            shd3.set(qn('w:color'), 'auto')
            shd3.set(qn('w:fill'), 'FFFED5')
            rPr3.append(shd3)

            # " Date "
            date_label = sig_para.add_run(' Date ')
            date_label.font.name = 'Arial'
            date_label.font.size = Pt(12)

            # Opening bracket for date (bold gold color)
            date_open = sig_para.add_run('[')
            date_open.font.name = 'Arial'
            date_open.font.size = Pt(12)
            date_open.font.bold = True
            date_open.font.color.rgb = BRACKET_COLOR
            rPr4 = date_open._element.get_or_add_rPr()
            shd4 = OxmlElement('w:shd')
            shd4.set(qn('w:val'), 'clear')
            shd4.set(qn('w:color'), 'auto')
            shd4.set(qn('w:fill'), 'FFFED5')
            rPr4.append(shd4)

            # Date content (cream highlighted)
            date_content = sig_para.add_run(sig_date)
            date_content.font.name = 'Arial'
            date_content.font.size = Pt(12)
            rPr5 = date_content._element.get_or_add_rPr()
            shd5 = OxmlElement('w:shd')
            shd5.set(qn('w:val'), 'clear')
            shd5.set(qn('w:color'), 'auto')
            shd5.set(qn('w:fill'), 'FFFED5')
            rPr5.append(shd5)

            # Closing bracket for date (bold gold color)
            date_close = sig_para.add_run(']')
            date_close.font.name = 'Arial'
            date_close.font.size = Pt(12)
            date_close.font.bold = True
            date_close.font.color.rgb = BRACKET_COLOR
            rPr6 = date_close._element.get_or_add_rPr()
            shd6 = OxmlElement('w:shd')
            shd6.set(qn('w:val'), 'clear')
            shd6.set(qn('w:color'), 'auto')
            shd6.set(qn('w:fill'), 'FFFED5')
            rPr6.append(shd6)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form A2 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        return {
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.toPlainText(),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.toPlainText(),
            "amhp_email": self.amhp_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.toPlainText(),
            "local_authority": self.local_authority.text(),
            "approved_by_same": self.approved_by_same.isChecked(),
            "approved_by_authority": self.approved_by_authority.text(),
            "nr_known": self.nr_known.isChecked(),
            "nr_option_a": self.nr_option_a.isChecked(),
            "nr_name": self.nr_name.text(),
            "nr_address": self.nr_address.toPlainText(),
            "nr_informed_yes": self.nr_informed_yes.isChecked(),
            "nr_unable": self.nr_unable.isChecked(),
            "last_seen_date": self.last_seen_date.date().toString("yyyy-MM-dd"),
            "no_acquaintance_reason": self.no_acquaintance_reason.toPlainText(),
            "signature_date": self.signature_date.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        if not state:
            return
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setPlainText(state.get("hospital_address", ""))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setPlainText(state.get("amhp_address", ""))
        self.amhp_email.setText(state.get("amhp_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setPlainText(state.get("patient_address", ""))
        self.local_authority.setText(state.get("local_authority", ""))
        if state.get("approved_by_same", True):
            self.approved_by_same.setChecked(True)
        else:
            self.approved_by_different.setChecked(True)
        self.approved_by_authority.setText(state.get("approved_by_authority", ""))
        if state.get("nr_known", True):
            self.nr_known.setChecked(True)
        else:
            self.nr_unknown.setChecked(True)
        if state.get("nr_option_a", True):
            self.nr_option_a.setChecked(True)
        else:
            self.nr_option_b.setChecked(True)
        self.nr_name.setText(state.get("nr_name", ""))
        self.nr_address.setPlainText(state.get("nr_address", ""))
        if state.get("nr_informed_yes", True):
            self.nr_informed_yes.setChecked(True)
        else:
            self.nr_informed_no.setChecked(True)
        if state.get("nr_unable", True):
            self.nr_unable.setChecked(True)
        else:
            self.nr_none.setChecked(True)
        if state.get("last_seen_date"):
            self.last_seen_date.setDate(QDate.fromString(state["last_seen_date"], "yyyy-MM-dd"))
        self.no_acquaintance_reason.setPlainText(state.get("no_acquaintance_reason", ""))
        if state.get("signature_date"):
            self.signature_date.setDate(QDate.fromString(state["signature_date"], "yyyy-MM-dd"))

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[A2Form] Set patient name: {patient_info['name']}")
