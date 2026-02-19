from __future__ import annotations

import os
import re
from datetime import datetime
from typing import List, Dict, Any

import fitz  # PyMuPDF
from docx import Document
import pandas as pd
from utils.report_detector import detect_report_type, is_blank_template, strip_form_headings

# ============================================================
#  DEBUG
# ============================================================

DEBUG = False


def _debug(msg: str):
    if DEBUG:
        print(f"[INGESTOR] {msg}")


# ============================================================
#  DATE INFERENCE
# ============================================================

DATE_PATTERNS = [
    r"\b(\d{2})[\/\-](\d{2})[\/\-](\d{4})\b",
    r"\b(\d{4})[\/\-](\d{2})[\/\-](\d{2})\b",
]


def infer_date_from_text(text: str) -> datetime | None:
    if not text:
        return None

    for pat in DATE_PATTERNS:
        m = re.search(pat, text)
        if not m:
            continue

        try:
            parts = m.groups()
            if len(parts[0]) == 4:
                return datetime(int(parts[0]), int(parts[1]), int(parts[2]))
            else:
                return datetime(int(parts[2]), int(parts[1]), int(parts[0]))
        except Exception:
            continue

    return None


def infer_date_from_filename(fname: str) -> datetime | None:
    return infer_date_from_text(fname)


def resolve_date(text: str, fname: str) -> datetime:
    return (
        infer_date_from_text(text)
        or infer_date_from_filename(fname)
        or datetime.today()
    )


# ============================================================
#  PDF INGESTION
# ============================================================

def ingest_pdf(path: str) -> List[Dict[str, Any]]:
    notes = []
    fname = os.path.basename(path)

    _debug(f"Reading PDF: {fname}")

    doc = fitz.open(path)
    full_text = []

    for i, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            full_text.append(f"\n=== PAGE {i + 1} ===\n{text}")

    combined = "\n".join(full_text).strip()

    # Skip blank templates
    if is_blank_template(combined):
        _debug(f"SKIPPED PDF (blank template): {fname}")
        return notes

    # Strip form headings from extracted text
    combined = strip_form_headings(combined)

    date = resolve_date(combined, fname)

    detected = detect_report_type(combined)
    report_type = detected["report_type"]

    notes.append({
        "date": date,
        "text": combined,
        "type": "uploaded_document",
        "report_type": report_type,
        "report_confidence": detected["confidence"],
        "source": {
            "filename": fname,
            "format": "pdf",
        }
    })

    _debug(f"PDF parsed | chars={len(combined)} | date={date.date()}")
    return notes


# ============================================================
#  DOCX INGESTION (PARAGRAPHS + TABLES)
# ============================================================

def ingest_docx(path: str) -> List[Dict[str, Any]]:
    notes = []
    fname = os.path.basename(path)

    _debug(f"Reading DOCX: {fname}")

    doc = Document(path)

    text_parts = []

    # -----------------------------
    # Paragraphs
    # -----------------------------
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            text_parts.append(p.text.strip())

    # -----------------------------
    # Tables (CRITICAL FOR RC REPORTS)
    # -----------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text_parts.append(cell_text)

    combined = "\n".join(text_parts).strip()

    # Skip blank templates
    if is_blank_template(combined):
        _debug(f"SKIPPED DOCX (blank template): {fname}")
        return notes

    # Strip form headings from extracted text
    combined = strip_form_headings(combined)

    date = resolve_date(combined, fname)

    detected = detect_report_type(combined)

    notes.append({
        "date": date,
        "text": combined,
        "type": "uploaded_document",
        "report_type": detected.get("report_type"),
        "report_confidence": detected.get("confidence"),
        "source": {
            "filename": fname,
            "format": "docx",
        }
    })

    _debug(f"DOCX parsed | chars={len(combined)} | date={date.date() if date else 'unknown'}")

    return notes



# ============================================================
#  EXCEL INGESTION
# ============================================================

def ingest_excel(path: str) -> List[Dict[str, Any]]:
    notes = []
    fname = os.path.basename(path)

    _debug(f"Reading Excel: {fname}")

    xls = pd.ExcelFile(path)

    for sheet in xls.sheet_names:
        df = xls.parse(sheet)

        for idx, row in df.iterrows():
            cells = [
                str(v).strip()
                for v in row.tolist()
                if pd.notna(v) and str(v).strip()
            ]

            if not cells:
                continue

            text = "\n".join(cells)

            # Skip blank templates
            if is_blank_template(text):
                continue

            date = resolve_date(text, fname)

            detected = detect_report_type(text)

            notes.append({
                "date": date,
                "text": text,
                "type": "uploaded_document",
                "report_type": detected.get("report_type"),
                "report_confidence": detected.get("confidence"),
                "source": {
                    "filename": fname,
                    "format": "excel",
                    "sheet": sheet,
                    "row": int(idx) + 1,
                }
            })

    _debug(f"Excel parsed | entries={len(notes)}")
    return notes

# ============================================================
#  MASTER INGESTOR
# ============================================================

def ingest_documents(paths: List[str], debug: bool = True) -> List[Dict[str, Any]]:
    global DEBUG
    DEBUG = debug

    all_notes: List[Dict[str, Any]] = []

    for path in paths:
        if not os.path.exists(path):
            _debug(f"File not found: {path}")
            continue

        ext = os.path.splitext(path)[1].lower()

        try:
            if ext == ".pdf":
                all_notes.extend(ingest_pdf(path))
            elif ext in (".docx",):
                all_notes.extend(ingest_docx(path))
            elif ext in (".xlsx", ".xls"):
                all_notes.extend(ingest_excel(path))
            elif ext in (".txt",):
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read().strip()

                # Skip blank templates
                if is_blank_template(text):
                    _debug(f"SKIPPED TXT (blank template): {os.path.basename(path)}")
                    continue

                # Strip form headings from extracted text
                text = strip_form_headings(text)

                detected = detect_report_type(text)
                report_type = detected["report_type"]

                date = resolve_date(text, os.path.basename(path))

                all_notes.append({
                    "date": date,
                    "text": text,
                    "type": "uploaded_document",
                    "report_type": report_type,
                    "report_confidence": detected["confidence"],
                    "source": {
                        "filename": os.path.basename(path),
                        "format": "txt",
                    }
                })

            else:
                _debug(f"Unsupported file type: {path}")

        except Exception as e:
            _debug(f"ERROR ingesting {path}: {e}")

    _debug(f"TOTAL NOTES EMITTED: {len(all_notes)}")
    return all_notes
