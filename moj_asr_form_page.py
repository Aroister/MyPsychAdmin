# ================================================================
#  MOJ ASR FORM PAGE — Annual Statutory Report for Restricted Patient
#  Ministry of Justice - Mental Health Casework Section
#  Based exactly on MOJ_ASR_TEMPLATE_2025.docx structure
# ================================================================

from __future__ import annotations
import sys
from datetime import datetime
from typing import Optional
from PySide6.QtCore import Qt, Signal, QDate, QEvent, QTimer
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTextBrowser,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QGridLayout, QRadioButton, QButtonGroup, QComboBox, QCompleter,
    QStyleFactory, QSlider, QSizePolicy, QColorDialog,
    QSplitter, QStackedWidget
)
from PySide6.QtGui import QColor, QFontDatabase

# Collapsible and Resizable sections for expandable panels
try:
    from background_history_popup import CollapsibleSection, ResizableSection
except ImportError:
    CollapsibleSection = None
    ResizableSection = None

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}

from utils.resource_path import resource_path

from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# ASR CARD WIDGET
# ================================================================
class ASRCardWidget(QFrame):
    """A clickable card for an ASR report section."""

    clicked = Signal(str)

    STYLE_NORMAL = """
        ASRCardWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
        }
        ASRCardWidget:hover {
            border-color: #991b1b;
            background: #fef2f2;
        }
        QLabel {
            background: transparent;
            border: none;
        }
        QRadioButton, QCheckBox {
            background: transparent;
            border: none;
        }
    """

    STYLE_SELECTED = """
        ASRCardWidget {
            background: #fee2e2;
            border: 2px solid #991b1b;
            border-left: 4px solid #7f1d1d;
            border-radius: 12px;
        }
        QLabel {
            background: transparent;
            border: none;
        }
        QRadioButton, QCheckBox {
            background: transparent;
            border: none;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
        """)
        header_row.addWidget(title_lbl)
        header_row.addStretch()

        layout.addLayout(header_row)

        # Editor
        self.editor = MyPsychAdminRichTextEditor()
        self.editor.setPlaceholderText("Click to edit...")
        self.editor.setReadOnly(False)
        self._editor_height = 180
        self.editor.setMinimumHeight(100)
        self.editor.setMaximumHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 16px;
                color: #374151;
            }
        """)
        layout.addWidget(self.editor)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.editor, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(12)
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #991b1b;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(60, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setMinimumHeight(self._editor_height)
                self.editor.setMaximumHeight(self._editor_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        return self._selected


# ================================================================
# MOJ ASR TOOLBAR
# ================================================================
class MOJASRToolbar(QWidget):
    """Toolbar for the MOJ ASR Form Page."""

    # Formatting signals
    set_font_family = Signal(str)
    set_font_size = Signal(int)
    toggle_bold = Signal()
    toggle_italic = Signal()
    toggle_underline = Signal()
    set_text_color = Signal(QColor)
    set_highlight_color = Signal(QColor)
    set_align_left = Signal()
    set_align_center = Signal()
    set_align_right = Signal()
    set_align_justify = Signal()
    bullet_list = Signal()
    numbered_list = Signal()
    indent = Signal()
    outdent = Signal()
    undo = Signal()
    redo = Signal()
    insert_date = Signal()

    # Action signals
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(80)
        self.setStyleSheet("""
            MOJASRToolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 500;
            }
            QToolButton:hover {
                background: rgba(0,0,0,0.08);
            }
            QComboBox {
                background: rgba(255,255,255,0.85);
                color: #333333;
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 16px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #e0e0e0;
            }
            QScrollArea {
                background: transparent;
                border: none;
            }
            QScrollBar:horizontal {
                height: 6px;
                background: transparent;
            }
            QScrollBar::handle:horizontal {
                background: rgba(0,0,0,0.2);
                border-radius: 3px;
                min-width: 30px;
            }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0px;
            }
        """)

        # Outer layout
        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        # Scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        # Container
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 8, 2)
        layout.setSpacing(10)

        # EXPORT DOCX
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(160, 38)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 17px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #1d4ed8; }
            QToolButton:pressed { background: #1e40af; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # FONT FAMILY
        self.font_combo = QComboBox()
        self.font_combo.setFixedWidth(160)
        families = QFontDatabase.families()
        if sys.platform == "win32":
            preferred = ["Segoe UI", "Calibri", "Cambria", "Arial", "Times New Roman"]
        else:
            preferred = ["Avenir Next", "Avenir", "SF Pro Text", "Helvetica Neue", "Helvetica"]
        added = set()
        for f in preferred:
            if f in families:
                self.font_combo.addItem(f)
                added.add(f)
        for f in families:
            if f not in added:
                self.font_combo.addItem(f)
        self.font_combo.currentTextChanged.connect(self.set_font_family.emit)
        layout.addWidget(self.font_combo)

        # FONT SIZE
        self.size_combo = QComboBox()
        self.size_combo.setFixedWidth(60)
        for sz in [8, 9, 10, 11, 12, 14, 16, 18, 20, 22]:
            self.size_combo.addItem(str(sz))
        self.size_combo.setCurrentText("12")
        self.size_combo.currentTextChanged.connect(lambda v: self.set_font_size.emit(int(v)))
        layout.addWidget(self.size_combo)

        # Button helper
        def btn(label, slot):
            b = QToolButton()
            b.setText(label)
            b.setMinimumWidth(36)
            b.clicked.connect(slot)
            return b

        # BASIC STYLES
        layout.addWidget(btn("B", self.toggle_bold.emit))
        layout.addWidget(btn("I", self.toggle_italic.emit))
        layout.addWidget(btn("U", self.toggle_underline.emit))

        # COLORS
        layout.addWidget(btn("A", self._choose_text_color))
        layout.addWidget(btn("🖍", self._choose_highlight_color))

        # ALIGNMENT
        layout.addWidget(btn("L", self.set_align_left.emit))
        layout.addWidget(btn("C", self.set_align_center.emit))
        layout.addWidget(btn("R", self.set_align_right.emit))
        layout.addWidget(btn("J", self.set_align_justify.emit))

        # LISTS / INDENTATION
        layout.addWidget(btn("•", self.bullet_list.emit))
        layout.addWidget(btn("1.", self.numbered_list.emit))
        layout.addWidget(btn("→", self.indent.emit))
        layout.addWidget(btn("←", self.outdent.emit))

        # UNDO / REDO
        layout.addWidget(btn("⟲", self.undo.emit))
        layout.addWidget(btn("⟳", self.redo.emit))

        # DATE INSERT
        layout.addWidget(btn("Date", self.insert_date.emit))

        layout.addStretch()

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)

    def _choose_text_color(self):
        col = QColorDialog.getColor(QColor("black"), self)
        if col.isValid():
            self.set_text_color.emit(col)

    def _choose_highlight_color(self):
        col = QColorDialog.getColor(QColor("yellow"), self)
        if col.isValid():
            self.set_highlight_color.emit(col)


# ================================================================
# MOJ ASR FORM PAGE
# ================================================================
class MOJASRFormPage(QWidget):
    """Page for completing MOJ Annual Statutory Report for Restricted Patients."""

    go_back = Signal()

    MHA_SECTIONS = [
        "S37/41",
        "S45a",
        "S47/49",
        "CPI - unfit to plead",
        "CPI - not guilty by reason of insanity",
        "DVCV"
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        # Track last generated text for each popup to preserve user additions
        self._last_generated_text = {}
        self._data_processed_id = None   # Signature of last processed extracted data (prevents reprocessing)
        self._notes_processed_id = None  # Signature of last processed notes (prevents reprocessing)
        self._setup_ui()
        self._prefill()
        self._connect_shared_store()

    def _connect_shared_store(self):
        """Connect to SharedDataStore for cross-report data sharing."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.notes_changed.connect(self._on_notes_changed)
            shared_store.extracted_data_changed.connect(self._on_extracted_data_changed)
            shared_store.patient_info_changed.connect(self._on_patient_info_changed)
            print("[MOJ-ASR] Connected to SharedDataStore signals (notes, extracted_data, patient_info)")

            # Check if there's already data in the store
            self._check_shared_store_for_existing_data()
        except Exception as e:
            print(f"[MOJ-ASR] Failed to connect to SharedDataStore: {e}")

    def _check_shared_store_for_existing_data(self):
        """Check SharedDataStore for existing data when page is created."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()

            # Check for existing patient info
            patient_info = shared_store.patient_info
            if patient_info and any(patient_info.values()):
                print(f"[MOJ-ASR] Found existing patient info in SharedDataStore")
                self._on_patient_info_changed(patient_info)

            # Check for existing notes
            notes = shared_store.notes
            if notes:
                print(f"[MOJ-ASR] Found {len(notes)} existing notes in SharedDataStore")
                self._on_notes_changed(notes)

            # Check for existing extracted data
            extracted_data = shared_store.extracted_data
            if extracted_data:
                print(f"[MOJ-ASR] Found existing extracted data in SharedDataStore")
                self._on_extracted_data_changed(extracted_data)
        except Exception as e:
            print(f"[MOJ-ASR] Error checking shared store: {e}")

    def _on_patient_info_changed(self, patient_info: dict):
        """Handle patient info updates from SharedDataStore."""
        if patient_info and any(patient_info.values()):
            print(f"[MOJ-ASR] Received patient info from SharedDataStore")
            # Fill patient fields if they exist
            if hasattr(self, 'popup_surname') and patient_info.get("name"):
                # Split name into parts
                name_parts = patient_info["name"].split()
                if len(name_parts) >= 2:
                    self.popup_surname.setText(name_parts[-1])
                    self.popup_forenames.setText(" ".join(name_parts[:-1]))

    def _on_notes_changed(self, notes: list):
        """Handle notes updates from SharedDataStore."""
        if notes:
            print(f"[MOJ-ASR] Received {len(notes)} notes from SharedDataStore")
            self._extracted_raw_notes = notes
            self._document_type = "notes"

    def _on_extracted_data_changed(self, data: dict):
        """Handle extracted data updates from SharedDataStore - trigger full processing."""
        if not data:
            return
        print(f"[MOJ-ASR] Received extracted data from SharedDataStore: {list(data.keys())}")
        # Trigger the same processing as local import
        self._on_data_extracted(data)

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        """Prefill RC details from my_details."""
        if hasattr(self, 'popup_rc_name') and self._my_details.get("full_name"):
            self.popup_rc_name.setText(self._my_details["full_name"])
        if hasattr(self, 'popup_rc_email') and self._my_details.get("email"):
            self.popup_rc_email.setText(self._my_details["email"])

    def set_notes(self, notes: list):
        """
        Set notes from shared data store.
        Called by MainWindow when notes are available from another section.
        """
        if not notes:
            return

        # Skip if these exact notes were already processed
        notes_sig = (len(notes), id(notes))
        if self._notes_processed_id == notes_sig:
            print(f"[MOJ-ASR] Skipping set_notes - notes already processed ({len(notes)} notes)")
            return
        self._notes_processed_id = notes_sig

        # Store raw notes at page level for use in sections
        self._extracted_raw_notes = notes
        self._document_type = "notes"

        # If data extractor exists, update its notes too
        if hasattr(self, '_data_extractor') and self._data_extractor:
            if hasattr(self._data_extractor, 'set_notes'):
                self._data_extractor.set_notes(notes)

        print(f"[MOJ-ASR] Received {len(notes)} notes from shared store")

    def _get_pronouns(self):
        """Get pronouns based on gender selection."""
        if hasattr(self, 'popup_gender_male') and self.popup_gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his", "self": "himself",
                    "has": "has", "is": "is", "engages": "engages", "attends": "attends", "lacks": "lacks", "suffers": "suffers", "does": "does", "sees": "sees"}
        elif hasattr(self, 'popup_gender_female') and self.popup_gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her", "self": "herself",
                    "has": "has", "is": "is", "engages": "engages", "attends": "attends", "lacks": "lacks", "suffers": "suffers", "does": "does", "sees": "sees"}
        else:
            return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their", "self": "themselves",
                    "has": "have", "is": "are", "engages": "engage", "attends": "attend", "lacks": "lack", "suffers": "suffer", "does": "do", "sees": "see"}

    # Section definitions for cards
    SECTIONS = [
        ("1. Patient Details", "patient_details"),
        ("2. Responsible Clinician", "rc_details"),
        ("3. Patient's Mental Disorder", "mental_disorder"),
        ("4. Attitude & Behaviour", "attitude_behaviour"),
        ("5. Addressing Issues", "addressing_issues"),
        ("6. Patient's Attitude", "patient_attitude"),
        ("7. Capacity Issues", "capacity"),
        ("8. Progress", "progress"),
        ("9. Managing Risk", "managing_risk"),
        ("10. How Risks Addressed", "risk_addressed"),
        ("11. Abscond / Escape", "abscond"),
        ("12. MAPPA", "mappa"),
        ("13. Victims", "victims"),
        ("14. Leave Report", "leave_report"),
        ("15. Additional Comments", "additional_comments"),
        ("16. Unfit to Plead", "unfit_to_plead"),
        ("Signature", "signature"),
    ]

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header with Clear Form button on right
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #991b1b; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("MOJ Annual Statutory Report — Restricted Patient")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Clear Form button in header (right side)
        clear_btn = QPushButton("Clear Form")
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setFixedSize(220, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover {
                background: #7f1d1d;
            }
            QPushButton:pressed {
                background: #450a0a;
            }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Full Toolbar with formatting options
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Data extractor instance (persists across views)
        self._data_extractor = None

        # Store raw notes for section 6 (like tribunal report does)
        self._extracted_raw_notes = []

        # Content area with splitter
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards scroll area
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("""
            QScrollArea {
                background: #f3f4f6;
                border: none;
            }
        """)
        self.main_splitter.addWidget(self.cards_holder)

        self.cards_root = QWidget()
        self.cards_root.setStyleSheet("background: #f3f4f6;")
        self.cards_layout = QVBoxLayout(self.cards_root)
        self.cards_layout.setContentsMargins(32, 24, 32, 24)
        self.cards_layout.setSpacing(16)
        self.cards_holder.setWidget(self.cards_root)

        # Right: Panel with popup stack
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(400)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame {
                background: rgba(245,245,245,0.98);
                border-left: 1px solid rgba(0,0,0,0.08);
            }
        """)
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 500])

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 22px;
            font-weight: 700;
            color: #991b1b;
            background: rgba(153, 27, 27, 0.1);
            padding: 8px 12px;
            border-radius: 8px;
        """)
        panel_layout.addWidget(self.panel_title)

        # Popup stack
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("background: white; border-radius: 8px;")
        self.popup_stack.setMinimumHeight(200)

        # Add a placeholder widget
        placeholder = QLabel("Click a section card to edit")
        placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        placeholder.setStyleSheet("font-size: 18px; color: #6b7280; background: white;")
        self.popup_stack.addWidget(placeholder)

        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        # Initialize cards and popups
        self.cards = {}
        self.popups = {}
        self.popup_previews = {}  # Preview labels for each popup
        self.popup_send_buttons = {}  # Send buttons for each popup
        self.popup_generators = {}  # Generator functions for each popup
        self._selected_card_key = None

        # Create all cards
        self._create_cards()

        # Build popup content (reuses existing section builder code)
        self._build_popups()

    def _register_active_editor(self, editor):
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        original_focus_in = editor.focusInEvent
        page = self
        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)
        editor.focusInEvent = focus_handler

    def _create_cards(self):
        """Create all section cards."""
        for title, key in self.SECTIONS:
            card = ASRCardWidget(title, key, parent=self.cards_root)
            card.clicked.connect(self._on_card_clicked)
            self.cards[key] = card
            self.cards_layout.addWidget(card)
            self._hook_editor_focus(card.editor)

        self.cards_layout.addStretch()

    def _on_card_clicked(self, key: str):
        """Handle card click - show appropriate popup."""
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        # Deselect previous
        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)

        # Select new
        self._selected_card_key = key
        if key in self.cards:
            self.cards[key].setSelected(True)

        # Show popup
        if key in self.popups:
            self.popup_stack.setCurrentWidget(self.popups[key])

    def _build_popups(self):
        """Build all section popups with their controls."""
        # Build each section's popup content
        self._build_popup_patient_details()
        self._build_popup_rc_details()
        self._build_popup_mental_disorder()
        self._build_popup_attitude_behaviour()
        self._build_popup_addressing_issues()
        self._build_popup_patient_attitude()
        self._build_popup_capacity()
        self._build_popup_progress()
        self._build_popup_managing_risk()
        self._build_popup_risk_addressed()
        self._build_popup_abscond()
        self._build_popup_mappa()
        self._build_popup_victims()
        self._build_popup_leave_report()
        self._build_popup_additional_comments()
        self._build_popup_unfit_to_plead()
        self._build_popup_signature()

    def _create_popup_container(self, key: str) -> tuple:
        """Create popup with input fields only (auto-syncs to card on change)."""
        # Main container - simple layout without preview section
        main_widget = QWidget()
        main_widget.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(0)

        # Scrollable input area
        input_scroll = QScrollArea()
        input_scroll.setWidgetResizable(True)
        input_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        input_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        input_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        input_container = QWidget()
        input_container.setStyleSheet("background: transparent;")
        input_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        input_layout = QVBoxLayout(input_container)
        input_layout.setContentsMargins(12, 12, 12, 12)
        input_layout.setSpacing(10)

        input_scroll.setWidget(input_container)
        main_layout.addWidget(input_scroll)

        # Add to popup stack
        self.popups[key] = main_widget
        self.popup_stack.addWidget(main_widget)

        return input_container, input_layout

    def _create_popup_container_with_imports(self, key: str) -> tuple:
        """Create popup with input and imported notes in a QSplitter (auto-syncs to card)."""
        # Main container
        main_widget = QWidget()
        main_widget.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(0)

        # Use QSplitter for input and imports sections
        splitter = QSplitter(Qt.Orientation.Vertical)
        splitter.setHandleWidth(6)
        splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(0,0,0,0.05), stop:0.5 rgba(0,0,0,0.15), stop:1 rgba(0,0,0,0.05));
                margin: 2px 60px;
                border-radius: 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(153,27,27,0.2), stop:0.5 rgba(153,27,27,0.5), stop:1 rgba(153,27,27,0.2));
            }
        """)

        # ============================================
        # INPUT SECTION (top pane)
        # ============================================
        input_widget = QWidget()
        input_widget.setStyleSheet("background: white;")
        input_widget_layout = QVBoxLayout(input_widget)
        input_widget_layout.setContentsMargins(0, 4, 0, 4)
        input_widget_layout.setSpacing(0)

        # Input header
        input_header = QFrame()
        input_header.setFixedHeight(28)
        input_header.setStyleSheet("""
            QFrame { background: rgba(153, 27, 27, 0.1); border: 1px solid rgba(153, 27, 27, 0.2); border-radius: 6px 6px 0 0; }
        """)
        input_header_layout = QHBoxLayout(input_header)
        input_header_layout.setContentsMargins(10, 4, 10, 4)
        input_header_label = QLabel("Input Fields")
        input_header_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        input_header_layout.addWidget(input_header_label)
        input_widget_layout.addWidget(input_header)

        input_scroll = QScrollArea()
        input_scroll.setWidgetResizable(True)
        input_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        input_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        input_scroll.setStyleSheet("QScrollArea { background: transparent; border: 1px solid rgba(153,27,27,0.2); border-top: none; border-radius: 0 0 6px 6px; }")

        input_container = QWidget()
        input_container.setStyleSheet("background: transparent;")
        input_layout = QVBoxLayout(input_container)
        input_layout.setContentsMargins(8, 8, 8, 8)
        input_layout.setSpacing(2)

        input_scroll.setWidget(input_container)
        input_widget_layout.addWidget(input_scroll)
        splitter.addWidget(input_widget)

        # ============================================
        # IMPORTED NOTES SECTION (bottom pane)
        # ============================================
        import_widget = QWidget()
        import_widget.setStyleSheet("background: white;")
        import_widget_layout = QVBoxLayout(import_widget)
        import_widget_layout.setContentsMargins(0, 4, 0, 0)
        import_widget_layout.setSpacing(0)

        # Import header
        import_header = QFrame()
        import_header.setFixedHeight(36)
        import_header.setStyleSheet("""
            QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 6px 6px 0 0; }
        """)
        import_header_layout = QHBoxLayout(import_header)
        import_header_layout.setContentsMargins(10, 4, 10, 4)
        import_header_layout.setSpacing(6)
        import_header_label = QLabel("Imported Data")
        import_header_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #806000; background: transparent; border: none;")
        import_header_layout.addWidget(import_header_label)

        # Scrollable container for category filter buttons (will be populated later)
        category_scroll = QScrollArea()
        category_scroll.setWidgetResizable(True)
        category_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        category_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        category_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        category_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; } QScrollBar:horizontal { height: 6px; }")
        category_scroll.setMaximumHeight(24)

        category_buttons_container = QWidget()
        category_buttons_container.setStyleSheet("background: transparent;")
        category_buttons_layout = QHBoxLayout(category_buttons_container)
        category_buttons_layout.setContentsMargins(0, 0, 0, 0)
        category_buttons_layout.setSpacing(4)
        category_scroll.setWidget(category_buttons_container)
        import_header_layout.addWidget(category_scroll, 1)
        import_widget_layout.addWidget(import_header)

        import_scroll = QScrollArea()
        import_scroll.setWidgetResizable(True)
        import_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        import_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        import_scroll.setStyleSheet("QScrollArea { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180,150,50,0.4); border-top: none; border-radius: 0 0 6px 6px; }")

        import_container = QWidget()
        import_container.setStyleSheet("background: transparent;")
        import_container.setMinimumWidth(0)
        import_container.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        import_layout = QVBoxLayout(import_container)
        import_layout.setContentsMargins(8, 8, 8, 8)
        import_layout.setSpacing(4)

        placeholder = QLabel("No imported data. Use Import File to upload data.")
        placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
        import_layout.addWidget(placeholder)

        import_scroll.setWidget(import_container)
        import_widget_layout.addWidget(import_scroll)
        splitter.addWidget(import_widget)

        # Set initial sizes (input: stretch, import: 150px)
        splitter.setSizes([400, 150])
        splitter.setStretchFactor(0, 1)  # Input stretches
        splitter.setStretchFactor(1, 0)  # Import doesn't stretch

        main_layout.addWidget(splitter)

        # Store import section references for later population
        setattr(self, f"popup_{key}_import_container", import_container)
        setattr(self, f"popup_{key}_import_layout", import_layout)
        setattr(self, f"popup_{key}_import_placeholder", placeholder)
        setattr(self, f"popup_{key}_imported_entries", [])
        setattr(self, f"popup_{key}_category_buttons_layout", category_buttons_layout)

        # Add to popup stack
        self.popups[key] = main_widget
        self.popup_stack.addWidget(main_widget)

        return input_container, input_layout

    def _create_4a_import_entry(self, text, date_str, preview_key,
                                 categories=None, category_colors=None,
                                 highlighted_html=None, filter_callback=None,
                                 parent_container=None):
        """Create a single import entry widget matching the 4a house style.

        Returns (entry_frame, checkbox, body_text) where body_text is a QTextEdit.
        """
        entry_frame = QFrame()
        entry_frame.setObjectName("entryFrame")
        entry_frame.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Maximum)
        entry_frame.setStyleSheet("""
            QFrame#entryFrame {
                background: rgba(255, 255, 255, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 8px;
                padding: 4px;
            }
        """)
        entry_layout = QVBoxLayout(entry_frame)
        entry_layout.setContentsMargins(10, 4, 16, 4)
        entry_layout.setSpacing(6)

        header_row = QHBoxLayout()
        header_row.setSpacing(8)

        # Toggle button (amber ▸/▾)
        toggle_btn = QPushButton("\u25b8")
        toggle_btn.setFixedSize(22, 22)
        toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        toggle_btn.setStyleSheet("""
            QPushButton {
                background: rgba(180, 150, 50, 0.2);
                border: none;
                border-radius: 4px;
                font-size: 15px;
                font-weight: bold;
                color: #806000;
            }
            QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
        """)
        header_row.addWidget(toggle_btn)

        # Date label with 📅
        date_display = f"\U0001f4c5 {date_str}" if date_str else "\U0001f4c5 No date"
        date_label = QLabel(date_display)
        date_label.setStyleSheet("""
            QLabel {
                font-size: 15px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)
        date_label.setCursor(Qt.CursorShape.PointingHandCursor)
        header_row.addWidget(date_label)

        # Category tag buttons (if provided)
        if categories and category_colors:
            for cat in categories:
                color = category_colors.get(cat, "#6b7280")
                if filter_callback:
                    tag = QPushButton(cat)
                    tag.setCursor(Qt.CursorShape.PointingHandCursor)
                    tag.setStyleSheet(f"""
                        QPushButton {{
                            font-size: 13px; font-weight: 600; color: white;
                            background: {color}; padding: 1px 4px; border-radius: 3px; border: none;
                        }}
                        QPushButton:hover {{ background: {color}; opacity: 0.8; }}
                    """)
                    tag.clicked.connect(lambda checked, c=cat: filter_callback(c))
                    header_row.addWidget(tag)
                else:
                    tag = QLabel(cat)
                    tag.setStyleSheet(f"font-size: 13px; font-weight: 600; color: white; background: {color}; padding: 1px 4px; border-radius: 3px;")
                    header_row.addWidget(tag)

        header_row.addStretch()

        # Checkbox (right side)
        cb = QCheckBox()
        cb.setProperty("full_text", text)
        cb.setFixedSize(20, 20)
        cb.setStyleSheet("""
            QCheckBox { background: transparent; margin-right: 4px; }
            QCheckBox::indicator { width: 18px; height: 18px; }
        """)
        cb.stateChanged.connect(lambda state: self._update_preview(preview_key))
        header_row.addWidget(cb)

        entry_layout.addLayout(header_row)

        # Body text (hidden QTextEdit, expandable)
        body_text = QTextEdit()
        if highlighted_html:
            body_text.setHtml(highlighted_html)
        else:
            body_text.setPlainText(text)
        body_text.setReadOnly(True)
        body_text.setFrameShape(QFrame.Shape.NoFrame)
        body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        body_text.setStyleSheet("""
            QTextEdit {
                font-size: 15px;
                color: #333;
                background: rgba(255, 248, 220, 0.5);
                border: none;
                padding: 8px;
                border-radius: 6px;
            }
        """)
        body_text.document().setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
        doc_height = body_text.document().size().height() + 20
        body_text.setFixedHeight(int(max(doc_height, 60)))
        body_text.setVisible(False)
        entry_layout.addWidget(body_text)

        # Toggle function
        container_ref = parent_container
        def make_toggle(btn, body, frame, container):
            def toggle():
                is_visible = body.isVisible()
                body.setVisible(not is_visible)
                btn.setText("\u25be" if not is_visible else "\u25b8")
                frame.updateGeometry()
                if container:
                    container.updateGeometry()
                    container.update()
            return toggle

        toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, container_ref)
        toggle_btn.clicked.connect(toggle_fn)
        date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

        return entry_frame, cb, body_text

    def _populate_popup_imports(self, key: str, entries: list):
        """Generic method to populate imported data panel for any popup."""
        import_layout = getattr(self, f"popup_{key}_import_layout", None)
        import_container = getattr(self, f"popup_{key}_import_container", None)
        if not import_layout or not import_container:
            return

        # Clear existing entries
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            setattr(self, f"popup_{key}_import_placeholder", placeholder)
            setattr(self, f"popup_{key}_imported_entries", [])
            return

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date = entry.get("date", "") or entry.get("datetime", "")
            if not text:
                continue

            date_str = str(date) if date else ""
            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=text,
                date_str=date_str,
                preview_key=key,
                parent_container=import_container,
            )
            import_layout.addWidget(entry_frame)

            entry_data = {"checkbox": cb, "text": text, "date": date}
            imported_entries.append(entry_data)

        setattr(self, f"popup_{key}_imported_entries", imported_entries)

    def _add_send_button(self, layout, key: str, generate_func):
        """Register generator function and trigger initial card sync."""
        # Store the generator function
        self.popup_generators[key] = generate_func

        # Initial update (auto-syncs to card)
        self._update_preview(key)

    def _update_preview(self, key: str):
        """Update the preview label and auto-sync to card, preserving user additions."""
        if key in self.popup_generators:
            try:
                new_generated_text = self.popup_generators[key]()
                # Update preview label if it exists
                if key in self.popup_previews:
                    self.popup_previews[key].setText(new_generated_text if new_generated_text else "(No content)")
                # Auto-sync to card, preserving user additions
                if key in self.cards:
                    current_text = self.cards[key].editor.toPlainText()
                    last_generated = self._last_generated_text.get(key, "")

                    if not current_text or current_text == last_generated:
                        # No user additions - just replace
                        self.cards[key].editor.setPlainText(new_generated_text if new_generated_text else "")
                    elif last_generated and last_generated in current_text:
                        # User added text after the generated content - preserve additions
                        # Find where the generated text ends and preserve everything after
                        idx = current_text.find(last_generated)
                        if idx == 0:
                            # Generated text is at the start - preserve text after it
                            user_additions = current_text[len(last_generated):]
                            self.cards[key].editor.setPlainText((new_generated_text or "") + user_additions)
                        else:
                            # Generated text is somewhere in the middle - preserve text before and after
                            before = current_text[:idx]
                            after = current_text[idx + len(last_generated):]
                            self.cards[key].editor.setPlainText(before + (new_generated_text or "") + after)
                    else:
                        # Last generated text not found - user heavily modified, append new text
                        # Don't replace to avoid losing user work
                        if new_generated_text and new_generated_text not in current_text:
                            # Only update if content meaningfully changed
                            self.cards[key].editor.setPlainText(new_generated_text if new_generated_text else "")

                    # Track the new generated text
                    self._last_generated_text[key] = new_generated_text or ""
            except Exception as e:
                if key in self.popup_previews:
                    self.popup_previews[key].setText(f"(Preview error: {e})")

    def _update_text_preserving_additions(self, text_widget, new_generated: str, tracking_key: str):
        """Update a text widget with new generated content while preserving user additions.

        Args:
            text_widget: The QTextEdit widget to update
            new_generated: The newly generated text
            tracking_key: Key to track the last generated text
        """
        current_text = text_widget.toPlainText()
        last_generated = self._last_generated_text.get(tracking_key, "")

        # Check for import markers that should be preserved
        import_marker = "\n\n--- Imported Notes ---\n"
        imported_section = ""
        if import_marker in current_text:
            imported_section = import_marker + current_text.split(import_marker)[1]
            current_text = current_text.split(import_marker)[0]
            if last_generated and import_marker in last_generated:
                last_generated = last_generated.split(import_marker)[0]

        if not current_text or current_text.strip() == last_generated.strip():
            # No user additions - just replace
            text_widget.setPlainText(new_generated + imported_section)
        elif last_generated and last_generated.strip() in current_text:
            # User added text - preserve additions
            last_gen_stripped = last_generated.strip()
            idx = current_text.find(last_gen_stripped)
            if idx == 0:
                user_additions = current_text[len(last_gen_stripped):]
                text_widget.setPlainText(new_generated + user_additions + imported_section)
            elif idx > 0:
                before = current_text[:idx]
                after = current_text[idx + len(last_gen_stripped):]
                text_widget.setPlainText(before + new_generated + after + imported_section)
            else:
                text_widget.setPlainText(new_generated + imported_section)
        else:
            # Last generated not found - user heavily modified
            # Just replace but preserve imported section
            text_widget.setPlainText(new_generated + imported_section)

        self._last_generated_text[tracking_key] = new_generated

    def _connect_preview_updates(self, key: str, widgets: list):
        """Connect widgets to trigger preview updates when changed."""
        for widget in widgets:
            if isinstance(widget, QLineEdit):
                widget.textChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QTextEdit):
                widget.textChanged.connect(lambda k=key: self._update_preview(k))
            elif isinstance(widget, QComboBox):
                widget.currentTextChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QCheckBox):
                widget.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QRadioButton):
                widget.toggled.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QDateEdit):
                widget.dateChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QSlider):
                widget.valueChanged.connect(lambda _, k=key: self._update_preview(k))

    def _send_to_card(self, key: str, text: str):
        """Send generated text to the card's editor."""
        if key in self.cards:
            current = self.cards[key].editor.toPlainText()
            if current:
                self.cards[key].editor.setPlainText(current + "\n\n" + text)
            else:
                self.cards[key].editor.setPlainText(text)

    # ================================================================
    # POPUP BUILDERS
    # ================================================================

    def _build_popup_patient_details(self):
        """Build patient details popup."""
        container, layout = self._create_popup_container("patient_details")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 17px; }"

        # Patient name
        name_lbl = QLabel("Patient Name:")
        name_lbl.setStyleSheet(label_style)
        layout.addWidget(name_lbl)
        self.popup_patient_name = QLineEdit()
        self.popup_patient_name.setStyleSheet(input_style)
        layout.addWidget(self.popup_patient_name)

        # Gender
        gender_lbl = QLabel("Gender:")
        gender_lbl.setStyleSheet(label_style)
        layout.addWidget(gender_lbl)
        gender_row = QHBoxLayout()
        self.popup_gender_group = QButtonGroup(self)
        self.popup_gender_male = QRadioButton("Male")
        self.popup_gender_male.setStyleSheet(radio_style)
        self.popup_gender_female = QRadioButton("Female")
        self.popup_gender_female.setStyleSheet(radio_style)
        self.popup_gender_group.addButton(self.popup_gender_male)
        self.popup_gender_group.addButton(self.popup_gender_female)
        gender_row.addWidget(self.popup_gender_male)
        gender_row.addWidget(self.popup_gender_female)
        gender_row.addStretch()
        layout.addLayout(gender_row)

        # DOB
        dob_lbl = QLabel("Date of Birth:")
        dob_lbl.setStyleSheet(label_style)
        layout.addWidget(dob_lbl)
        self.popup_dob = QDateEdit()
        self.popup_dob.setCalendarPopup(True)
        self.popup_dob.setDisplayFormat("dd/MM/yyyy")
        self.popup_dob.setStyleSheet("""
            QDateEdit {
                background-color: white;
                color: #1f2937;
                padding: 6px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                font-size: 17px;
                selection-background-color: #3b82f6;
                selection-color: white;
            }
            QDateEdit::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                background-color: #f9fafb;
                border-left: 1px solid #d1d5db;
                width: 20px;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
            }
            QDateEdit::down-arrow {
                image: none;
                width: 10px;
                height: 10px;
            }
            QDateEdit QAbstractItemView {
                background-color: white;
                color: #1f2937;
                selection-background-color: #3b82f6;
                selection-color: white;
            }
            QCalendarWidget {
                background-color: white;
            }
            QCalendarWidget QWidget {
                background-color: white;
                color: #1f2937;
            }
            QCalendarWidget QToolButton {
                background-color: white;
                color: #1f2937;
            }
            QCalendarWidget QSpinBox {
                background-color: white;
                color: #1f2937;
            }
            QCalendarWidget QMenu {
                background-color: white;
                color: #1f2937;
            }
        """)
        layout.addWidget(self.popup_dob)

        # NHS Number
        nhs_lbl = QLabel("NHS Number:")
        nhs_lbl.setStyleSheet(label_style)
        layout.addWidget(nhs_lbl)
        self.popup_nhs = QLineEdit()
        self.popup_nhs.setStyleSheet(input_style)
        layout.addWidget(self.popup_nhs)

        # Hospital
        hosp_lbl = QLabel("Hospital:")
        hosp_lbl.setStyleSheet(label_style)
        layout.addWidget(hosp_lbl)
        self.popup_hospital = QLineEdit()
        self.popup_hospital.setStyleSheet(input_style)
        layout.addWidget(self.popup_hospital)

        # MHA Section
        mha_lbl = QLabel("MHA Section:")
        mha_lbl.setStyleSheet(label_style)
        layout.addWidget(mha_lbl)
        self.popup_mha_section = QComboBox()
        self.popup_mha_section.addItems(self.MHA_SECTIONS)
        self.popup_mha_section.setStyleSheet("QComboBox { font-size: 17px; } QComboBox QAbstractItemView { font-size: 17px; }")
        layout.addWidget(self.popup_mha_section)

        # MHCS Reference No.
        mhcs_lbl = QLabel("MHCS Reference No.:")
        mhcs_lbl.setStyleSheet(label_style)
        layout.addWidget(mhcs_lbl)
        self.popup_mhcs_ref = QLineEdit()
        self.popup_mhcs_ref.setStyleSheet(input_style)
        layout.addWidget(self.popup_mhcs_ref)

        # MHA Section Date
        mha_date_lbl = QLabel("MHA Section Date:")
        mha_date_lbl.setStyleSheet(label_style)
        layout.addWidget(mha_date_lbl)
        self.popup_mha_section_date = QDateEdit()
        self.popup_mha_section_date.setCalendarPopup(True)
        self.popup_mha_section_date.setDisplayFormat("dd/MM/yyyy")
        self.popup_mha_section_date.setStyleSheet("""
            QDateEdit {
                background-color: white;
                color: #1f2937;
                padding: 6px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                font-size: 17px;
                selection-background-color: #3b82f6;
                selection-color: white;
            }
            QDateEdit::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                background-color: #f9fafb;
                border-left: 1px solid #d1d5db;
                width: 20px;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
            }
            QDateEdit::down-arrow {
                image: none;
                width: 10px;
                height: 10px;
            }
            QDateEdit QAbstractItemView {
                background-color: white;
                color: #1f2937;
                selection-background-color: #3b82f6;
                selection-color: white;
            }
            QCalendarWidget {
                background-color: white;
            }
            QCalendarWidget QWidget {
                background-color: white;
                color: #1f2937;
            }
            QCalendarWidget QToolButton {
                background-color: white;
                color: #1f2937;
            }
            QCalendarWidget QSpinBox {
                background-color: white;
                color: #1f2937;
            }
            QCalendarWidget QMenu {
                background-color: white;
                color: #1f2937;
            }
        """)
        layout.addWidget(self.popup_mha_section_date)

        # Other Detention Authorities
        other_det_lbl = QLabel("Other Detention Authorities:")
        other_det_lbl.setStyleSheet(label_style)
        layout.addWidget(other_det_lbl)
        self.popup_other_detention = QLineEdit()
        self.popup_other_detention.setStyleSheet(input_style)
        self.popup_other_detention.setPlaceholderText("e.g., Notional 37, Immigration hold")
        layout.addWidget(self.popup_other_detention)

        layout.addStretch()
        self._add_send_button(layout, "patient_details", self._generate_patient_details)
        self._connect_preview_updates("patient_details", [
            self.popup_patient_name, self.popup_gender_male, self.popup_gender_female,
            self.popup_dob, self.popup_nhs, self.popup_hospital, self.popup_mha_section,
            self.popup_mhcs_ref, self.popup_mha_section_date, self.popup_other_detention
        ])

    def _generate_patient_details(self) -> str:
        name = self.popup_patient_name.text() or "[Patient Name]"
        gender = "Male" if self.popup_gender_male.isChecked() else "Female" if self.popup_gender_female.isChecked() else "[Gender]"
        dob = self.popup_dob.date().toString("dd/MM/yyyy")
        nhs = self.popup_nhs.text() or "[NHS Number]"
        hospital = self.popup_hospital.text() or "[Hospital]"
        mha = self.popup_mha_section.currentText()
        mhcs_ref = self.popup_mhcs_ref.text() if hasattr(self, 'popup_mhcs_ref') else ""
        mha_section_date = self.popup_mha_section_date.date().toString("dd/MM/yyyy") if hasattr(self, 'popup_mha_section_date') else ""
        other_detention = self.popup_other_detention.text() if hasattr(self, 'popup_other_detention') and self.popup_other_detention.text() else "Nil"
        return f"Patient: {name}\nGender: {gender}\nDOB: {dob}\nNHS: {nhs}\nHospital: {hospital}\nMHA Section: {mha}\nMHCS Reference No.: {mhcs_ref}\nMHA Section Date: {mha_section_date}\nOther Detention Authorities: {other_detention}"

    def _build_popup_rc_details(self):
        """Build RC details popup."""
        container, layout = self._create_popup_container("rc_details")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        rc_name_lbl = QLabel("RC Name:")
        rc_name_lbl.setStyleSheet(label_style)
        layout.addWidget(rc_name_lbl)
        self.popup_rc_name = QLineEdit()
        self.popup_rc_name.setStyleSheet(input_style)
        if self._my_details.get("full_name"):
            self.popup_rc_name.setText(self._my_details["full_name"])
        layout.addWidget(self.popup_rc_name)

        rc_email_lbl = QLabel("RC Email:")
        rc_email_lbl.setStyleSheet(label_style)
        layout.addWidget(rc_email_lbl)
        self.popup_rc_email = QLineEdit()
        self.popup_rc_email.setStyleSheet(input_style)
        if self._my_details.get("email"):
            self.popup_rc_email.setText(self._my_details["email"])
        layout.addWidget(self.popup_rc_email)

        # Job Title
        job_lbl = QLabel("Job Title:")
        job_lbl.setStyleSheet(label_style)
        layout.addWidget(job_lbl)
        self.popup_rc_job_title = QLineEdit()
        self.popup_rc_job_title.setStyleSheet(input_style)
        self.popup_rc_job_title.setPlaceholderText("e.g., Consultant Forensic Psychiatrist")
        if self._my_details.get("role_title"):
            self.popup_rc_job_title.setText(self._my_details["role_title"])
        layout.addWidget(self.popup_rc_job_title)

        # Telephone No.
        phone_lbl = QLabel("Telephone No.:")
        phone_lbl.setStyleSheet(label_style)
        layout.addWidget(phone_lbl)
        self.popup_rc_phone = QLineEdit()
        self.popup_rc_phone.setStyleSheet(input_style)
        if self._my_details.get("phone"):
            self.popup_rc_phone.setText(self._my_details["phone"])
        layout.addWidget(self.popup_rc_phone)

        # MHA Office Email
        mha_email_lbl = QLabel("Mental Health Act Office Email:")
        mha_email_lbl.setStyleSheet(label_style)
        layout.addWidget(mha_email_lbl)
        self.popup_mha_office_email = QLineEdit()
        self.popup_mha_office_email.setStyleSheet(input_style)
        layout.addWidget(self.popup_mha_office_email)

        layout.addStretch()
        self._add_send_button(layout, "rc_details", self._generate_rc_details)
        self._connect_preview_updates("rc_details", [
            self.popup_rc_name, self.popup_rc_email, self.popup_rc_job_title,
            self.popup_rc_phone, self.popup_mha_office_email
        ])

    def _generate_rc_details(self) -> str:
        name = self.popup_rc_name.text() or "[RC Name]"
        email = self.popup_rc_email.text() or "[RC Email]"
        job_title = self.popup_rc_job_title.text() if hasattr(self, 'popup_rc_job_title') else ""
        phone = self.popup_rc_phone.text() if hasattr(self, 'popup_rc_phone') else ""
        mha_office_email = self.popup_mha_office_email.text() if hasattr(self, 'popup_mha_office_email') else ""

        parts = [f"Responsible Clinician: {name}"]
        parts.append(f"Job Title: {job_title}")
        parts.append(f"Telephone No.: {phone}")
        parts.append(f"Email: {email}")
        parts.append(f"MHA Office Email: {mha_office_email}")
        return "\n".join(parts)

    def _build_popup_mental_disorder(self):
        """Build mental disorder popup with 3 grouped ICD-10 dropdowns matching iOS style."""
        container, layout = self._create_popup_container("mental_disorder")

        icd_header = QLabel("<b>ICD-10 Diagnoses:</b>")
        icd_header.setStyleSheet("font-size: 18px;")
        layout.addWidget(icd_header)
        layout.addSpacing(8)

        # Grouped ICD-10 diagnoses matching iOS structure
        self.ICD10_GROUPED = [
            ("Schizophrenia & Psychosis", [
                ("F20.0", "Paranoid schizophrenia"),
                ("F20.1", "Hebephrenic schizophrenia"),
                ("F20.2", "Catatonic schizophrenia"),
                ("F20.3", "Undifferentiated schizophrenia"),
                ("F20.5", "Residual schizophrenia"),
                ("F20.6", "Simple schizophrenia"),
                ("F20.9", "Schizophrenia, unspecified"),
                ("F21", "Schizotypal disorder"),
                ("F22", "Persistent delusional disorders"),
                ("F23", "Acute and transient psychotic disorders"),
                ("F25.0", "Schizoaffective disorder, manic type"),
                ("F25.1", "Schizoaffective disorder, depressive type"),
                ("F25.2", "Schizoaffective disorder, mixed type"),
                ("F25.9", "Schizoaffective disorder, unspecified"),
                ("F29", "Unspecified nonorganic psychosis"),
            ]),
            ("Mood Disorders - Bipolar", [
                ("F30.0", "Hypomania"),
                ("F30.1", "Mania without psychotic symptoms"),
                ("F30.2", "Mania with psychotic symptoms"),
                ("F31.0", "Bipolar disorder, current episode hypomanic"),
                ("F31.1", "Bipolar disorder, current episode manic without psychosis"),
                ("F31.2", "Bipolar disorder, current episode manic with psychosis"),
                ("F31.3", "Bipolar disorder, current episode mild/moderate depression"),
                ("F31.4", "Bipolar disorder, current episode severe depression without psychosis"),
                ("F31.5", "Bipolar disorder, current episode severe depression with psychosis"),
                ("F31.6", "Bipolar disorder, current episode mixed"),
                ("F31.7", "Bipolar disorder, currently in remission"),
                ("F31.9", "Bipolar disorder, unspecified"),
            ]),
            ("Mood Disorders - Depression", [
                ("F32.0", "Mild depressive episode"),
                ("F32.1", "Moderate depressive episode"),
                ("F32.2", "Severe depressive episode without psychosis"),
                ("F32.3", "Severe depressive episode with psychosis"),
                ("F32.9", "Depressive episode, unspecified"),
                ("F33.0", "Recurrent depression, current episode mild"),
                ("F33.1", "Recurrent depression, current episode moderate"),
                ("F33.2", "Recurrent depression, current episode severe without psychosis"),
                ("F33.3", "Recurrent depression, current episode severe with psychosis"),
                ("F33.9", "Recurrent depressive disorder, unspecified"),
            ]),
            ("Anxiety Disorders", [
                ("F40.0", "Agoraphobia"),
                ("F40.1", "Social phobias"),
                ("F40.2", "Specific (isolated) phobias"),
                ("F41.0", "Panic disorder"),
                ("F41.1", "Generalized anxiety disorder"),
                ("F41.2", "Mixed anxiety and depressive disorder"),
                ("F42", "Obsessive-compulsive disorder"),
                ("F43.0", "Acute stress reaction"),
                ("F43.1", "Post-traumatic stress disorder"),
                ("F43.2", "Adjustment disorders"),
            ]),
            ("Eating Disorders", [
                ("F50.0", "Anorexia nervosa"),
                ("F50.2", "Bulimia nervosa"),
            ]),
            ("Personality Disorders", [
                ("F60.0", "Paranoid personality disorder"),
                ("F60.1", "Schizoid personality disorder"),
                ("F60.2", "Dissocial personality disorder"),
                ("F60.3", "Emotionally unstable personality disorder"),
                ("F60.4", "Histrionic personality disorder"),
                ("F60.5", "Anankastic personality disorder"),
                ("F60.6", "Anxious personality disorder"),
                ("F60.7", "Dependent personality disorder"),
                ("F60.9", "Personality disorder, unspecified"),
            ]),
            ("Intellectual Disability", [
                ("F70", "Mild intellectual disability"),
                ("F71", "Moderate intellectual disability"),
                ("F72", "Severe intellectual disability"),
                ("F79", "Unspecified intellectual disability"),
            ]),
            ("Organic Disorders", [
                ("F00", "Dementia in Alzheimer's disease"),
                ("F01", "Vascular dementia"),
                ("F03", "Unspecified dementia"),
                ("F05", "Delirium"),
                ("F06", "Other mental disorders due to brain damage"),
            ]),
            ("Substance Use Disorders", [
                ("F10", "Mental disorders due to alcohol"),
                ("F11", "Mental disorders due to opioids"),
                ("F12", "Mental disorders due to cannabinoids"),
                ("F14", "Mental disorders due to cocaine"),
                ("F15", "Mental disorders due to stimulants"),
                ("F19", "Mental disorders due to multiple drug use"),
            ]),
        ]

        # Create 3 ICD-10 dropdown combos with grouped items
        self.popup_dx_combos = []
        combo_style = """
            QComboBox {
                padding: 8px 32px 8px 8px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                font-size: 16px;
                background: white;
                color: #111827;
            }
            QComboBox QLineEdit {
                background: white;
                color: #111827;
                font-size: 16px;
                padding: 0px;
                border: none;
            }
            QComboBox:hover {
                border-color: #991b1b;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 28px;
                border-left: 1px solid #d1d5db;
                background: #f9fafb;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 6px solid #374151;
            }
            QComboBox::down-arrow:hover {
                border-top-color: #991b1b;
            }
            QComboBox QAbstractItemView {
                background: white;
                selection-background-color: #fee2e2;
                outline: none;
                font-size: 16px;
                color: #111827;
            }
        """

        labels = ["Primary Diagnosis:", "Secondary Diagnosis:", "Third Diagnosis:"]
        for i in range(3):
            lbl = QLabel(labels[i])
            lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
            layout.addWidget(lbl)

            combo = QComboBox()
            combo.setEditable(True)
            combo.setStyleSheet(combo_style)
            combo.setMinimumHeight(36)
            combo.setSizeAdjustPolicy(QComboBox.SizeAdjustPolicy.AdjustToMinimumContentsLengthWithIcon)
            combo.setMinimumContentsLength(30)
            combo.addItem("Select diagnosis...")  # Empty first option

            # Add grouped ICD-10 codes with category headers
            all_items = ["Select diagnosis..."]
            for group_name, diagnoses in self.ICD10_GROUPED:
                # Add category header (disabled, styled differently)
                combo.addItem(f"── {group_name} ──")
                idx = combo.count() - 1
                combo.model().item(idx).setEnabled(False)
                combo.model().item(idx).setData(QColor("#6b7280"), Qt.ItemDataRole.ForegroundRole)
                font = combo.model().item(idx).font()
                font.setBold(True)
                combo.model().item(idx).setFont(font)

                # Add diagnoses in this category
                for code, name in diagnoses:
                    display_text = f"{code} {name}"
                    combo.addItem(display_text, code)  # Store code as data
                    all_items.append(display_text)

            # Enable filtering/completion
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            completer = QCompleter(all_items)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            combo.setCompleter(completer)

            layout.addWidget(combo)
            self.popup_dx_combos.append(combo)
            layout.addSpacing(8)

        layout.addSpacing(12)
        clin_desc_lbl = QLabel("<b>Clinical Description:</b>")
        clin_desc_lbl.setStyleSheet("font-size: 18px;")
        layout.addWidget(clin_desc_lbl)
        self.popup_disorder_desc = QTextEdit()
        self.popup_disorder_desc.setMaximumHeight(150)
        self.popup_disorder_desc.setPlaceholderText("Additional clinical details about the mental disorder...")
        self.popup_disorder_desc.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 8px; font-size: 18px;")
        layout.addWidget(self.popup_disorder_desc)

        layout.addStretch()
        self._add_send_button(layout, "mental_disorder", self._generate_mental_disorder)
        self._connect_preview_updates("mental_disorder", self.popup_dx_combos + [self.popup_disorder_desc])

    def _generate_mental_disorder(self) -> str:
        """Generate mental disorder text from ICD-10 selections in narrative format."""
        diagnoses = []
        for combo in self.popup_dx_combos:
            text = combo.currentText().strip()
            # Skip empty or placeholder selections
            if text and text != "Select diagnosis..." and not text.startswith("──"):
                diagnoses.append(text)

        desc = self.popup_disorder_desc.toPlainText().strip()

        if not diagnoses:
            return desc if desc else "[No diagnoses selected]"

        # Build narrative format
        if len(diagnoses) == 1:
            narrative = f"Main diagnosis is {diagnoses[0]}."
        elif len(diagnoses) == 2:
            narrative = f"Main diagnosis is {diagnoses[0]}. There is a second diagnosis of {diagnoses[1]}."
        else:
            narrative = f"Main diagnosis is {diagnoses[0]}. There is a second diagnosis of {diagnoses[1]}; and as a third - {diagnoses[2]}."

        if desc:
            narrative += f"\n\n{desc}"

        return narrative

    def _build_popup_attitude_behaviour(self):
        """Build attitude/behaviour popup with behaviour categories, additional notes, and separate imported data section."""
        # Custom structure: Preview + Input Fields (collapsible) + Imported Notes (separate)
        key = "attitude_behaviour"

        # Main container with scroll area to prevent overlap
        main_widget = QWidget()
        main_widget.setStyleSheet("background: white;")
        outer_layout = QVBoxLayout(main_widget)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        # Scroll area for all content
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setStyleSheet("QScrollArea { background: white; border: none; }")

        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(scroll_content)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(12)

        # ============================================
        # INPUT FIELDS SECTION (auto-syncs to card)
        # ============================================
        if CollapsibleSection:
            input_section = CollapsibleSection("Input Fields", start_collapsed=False)
            input_section.set_content_height(280)
            input_section._min_height = 100
            input_section._max_height = 450
            input_section.set_header_style("""
                QFrame { background: rgba(153, 27, 27, 0.1); border: 1px solid rgba(153, 27, 27, 0.2); border-radius: 6px 6px 0 0; }
            """)
        else:
            input_section = QFrame()
            input_section.setStyleSheet("QFrame { background: transparent; }")

        input_scroll = QScrollArea()
        input_scroll.setWidgetResizable(True)
        input_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        input_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        input_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        input_container = QWidget()
        input_container.setStyleSheet("background: transparent;")
        layout = QVBoxLayout(input_container)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(10)

        # --- Behaviour Categories ---
        layout.addWidget(QLabel("<b>Behaviour Categories (last 12 months):</b>"))
        layout.addSpacing(4)

        self.popup_behaviour_items = {}
        behaviour_categories = [
            ("verbal_physical", "Verbal/physical aggression or violence", "no verbal or physical aggression"),
            ("substance_abuse", "Substance abuse", "no substance abuse"),
            ("self_harm", "Self-harm", "no self-harm"),
            ("fire_setting", "Fire-setting", "no fire-setting"),
            ("intimidation", "Intimidation/threats", "no intimidation or threats"),
            ("secretive", "Secretive/dishonest/manipulative behaviour", "no secretive or manipulative behaviour"),
            ("subversive", "Subversive behaviour", "no subversive behaviour"),
            ("sexually_disinhibited", "Sexually disinhibited/inappropriate behaviour", "no sexually disinhibited behaviour"),
            ("extremist", "Extremist/Terrorist risk ideology/behaviour", "no extremist behaviour"),
            ("seclusion", "Periods of seclusion", "no periods of seclusion"),
        ]

        grid_widget = QWidget()
        grid_widget.setStyleSheet("background: #f9fafb; border-radius: 8px;")
        grid_layout = QGridLayout(grid_widget)
        grid_layout.setContentsMargins(12, 12, 12, 12)
        grid_layout.setSpacing(8)

        row = 0
        for bkey, label, negative_text in behaviour_categories:
            lbl = QLabel(label)
            lbl.setStyleSheet("font-size: 16px; color: #374151; background: transparent;")
            lbl.setWordWrap(True)
            grid_layout.addWidget(lbl, row, 0)

            btn_widget = QWidget()
            btn_widget.setStyleSheet("background: transparent;")
            btn_layout = QHBoxLayout(btn_widget)
            btn_layout.setContentsMargins(0, 0, 0, 0)
            btn_layout.setSpacing(8)

            yes_rb = QRadioButton("Yes")
            yes_rb.setStyleSheet("""
                QRadioButton {
                    font-size: 15px;
                    color: #059669;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
            no_rb = QRadioButton("No")
            no_rb.setStyleSheet("""
                QRadioButton {
                    font-size: 15px;
                    color: #dc2626;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)

            btn_group = QButtonGroup(btn_widget)
            btn_group.addButton(yes_rb, 1)
            btn_group.addButton(no_rb, 0)

            btn_layout.addWidget(yes_rb)
            btn_layout.addWidget(no_rb)
            btn_layout.addStretch()
            grid_layout.addWidget(btn_widget, row, 1)

            details_input = QLineEdit()
            details_input.setPlaceholderText("Details...")
            details_input.setStyleSheet("font-size: 15px; padding: 4px 8px; border: 1px solid #d1d5db; border-radius: 4px; background: white;")
            details_input.hide()
            grid_layout.addWidget(details_input, row, 2)

            yes_rb.toggled.connect(lambda checked, d=details_input: d.setVisible(checked))

            self.popup_behaviour_items[bkey] = {
                "yes": yes_rb, "no": no_rb, "details": details_input,
                "label": label, "negative": negative_text
            }
            row += 1

        layout.addWidget(grid_widget)

        # --- Additional Notes (inside input section) ---
        layout.addSpacing(8)
        layout.addWidget(QLabel("<b>Additional Notes:</b>"))
        self.popup_attitude_notes = QTextEdit()
        self.popup_attitude_notes.setMaximumHeight(70)
        self.popup_attitude_notes.setPlaceholderText("Additional observations about attitude and behaviour...")
        self.popup_attitude_notes.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 8px;")
        layout.addWidget(self.popup_attitude_notes)

        input_scroll.setWidget(input_container)

        if CollapsibleSection:
            input_section.set_content(input_scroll)
        else:
            fallback_input_layout = QVBoxLayout(input_section)
            fallback_input_layout.setContentsMargins(0, 0, 0, 0)
            fallback_input_layout.addWidget(input_scroll)

        main_layout.addWidget(input_section)
        main_layout.addSpacing(8)

        # ============================================
        # IMPORTED NOTES SECTION (separate, at bottom)
        # ============================================
        if CollapsibleSection:
            import_section = CollapsibleSection("Imported Data", start_collapsed=True)
            import_section.set_content_height(150)
            import_section._min_height = 80
            import_section._max_height = 300
            import_section.set_header_style("""
                QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 6px 6px 0 0; }
            """)
            import_section.set_title_style("""
                QLabel { font-size: 17px; font-weight: 600; color: #806000; background: transparent; border: none; }
            """)
        else:
            import_section = QFrame()
            import_section.setStyleSheet("QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 8px; }")

        import_scroll = QScrollArea()
        import_scroll.setWidgetResizable(True)
        import_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        import_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        import_scroll.setStyleSheet("QScrollArea { background: rgba(255, 248, 220, 0.95); border: none; }")

        self.popup_behaviour_import_container = QWidget()
        self.popup_behaviour_import_container.setStyleSheet("background: transparent;")
        self.popup_behaviour_import_layout = QVBoxLayout(self.popup_behaviour_import_container)
        self.popup_behaviour_import_layout.setContentsMargins(8, 8, 8, 8)
        self.popup_behaviour_import_layout.setSpacing(4)

        self.popup_behaviour_import_placeholder = QLabel("No imported data. Use Import File to upload data.")
        self.popup_behaviour_import_placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
        self.popup_behaviour_import_layout.addWidget(self.popup_behaviour_import_placeholder)

        import_scroll.setWidget(self.popup_behaviour_import_container)

        if CollapsibleSection:
            import_section.set_content(import_scroll)
        else:
            fallback_import_layout = QVBoxLayout(import_section)
            fallback_import_layout.setContentsMargins(8, 8, 8, 8)
            fallback_import_layout.addWidget(import_scroll)

        main_layout.addWidget(import_section)
        main_layout.addStretch()

        self.popup_behaviour_imported_entries = []

        # Wire up scroll area
        main_scroll.setWidget(scroll_content)
        outer_layout.addWidget(main_scroll)

        # Add to popup stack
        self.popups[key] = main_widget
        self.popup_stack.addWidget(main_widget)

        # Connect send button and preview updates
        self._add_send_button(layout, key, self._generate_attitude_behaviour)
        behaviour_widgets = [self.popup_attitude_notes]
        for cat_widgets in self.popup_behaviour_items.values():
            behaviour_widgets.extend([cat_widgets["yes"], cat_widgets["no"], cat_widgets["details"]])
        self._connect_preview_updates(key, behaviour_widgets)

    def populate_popup_behaviour_imports(self, entries: list):
        """Populate the imported data panel in the behaviour popup with categorized entries.

        Filters to 12-month window from latest entry and categorizes by: compliance, admissions, insight, engagement.
        Relevant labeled notes appear at the top.
        """
        from datetime import datetime, timedelta

        print(f"[MOJ-ASR] Section 4 popup: populate_popup_behaviour_imports called with {len(entries) if entries else 0} entries")

        # Check if UI elements exist
        if not hasattr(self, 'popup_behaviour_import_layout') or not self.popup_behaviour_import_layout:
            print("[MOJ-ASR] Section 4 popup: popup_behaviour_import_layout not available")
            return

        # Clear existing entries
        while self.popup_behaviour_import_layout.count():
            item = self.popup_behaviour_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_behaviour_imported_entries = []

        if not entries:
            self.popup_behaviour_import_placeholder = QLabel("No imported data. Use Import File to upload data.")
            self.popup_behaviour_import_placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            self.popup_behaviour_import_layout.addWidget(self.popup_behaviour_import_placeholder)
            print("[MOJ-ASR] Section 4 popup: No entries to display")
            return

        # Category keywords for behaviour/attitude relevance
        # Enhanced with comprehensive risk patterns from violence_risk_analyzer
        CATEGORIES = {
            "Compliance": [
                "complian", "adherent", "taking medication", "treatment adherence",
                "concordan", "engagement with treatment"
            ],
            "Noncompliance": [
                "non compliance", "noncompliance", "non-compliance",
                "noncompliant", "non-compliant", "not compliant",
                "non-adherent", "nonadherent", "discordant",
                "refused medication", "refusing medication", "declined medication",
                "refused meds", "refusing meds", "declined meds", "not taking medication",
                "stopped taking", "refused depot", "declined depot",
                "refused clozapine", "declined clozapine",
                "refused ot", "declined ot", "refused occupational therapy", "declined occupational therapy",
                "refused psychology", "declined psychology", "refused psychological", "declined psychological",
                "refused therapy", "declined therapy",
                "refused treatment", "refusing treatment", "declined treatment",
                "refused to engage", "refusing to engage", "declined to engage",
                "refused to attend", "refusing to attend", "declined to attend",
                "refused nursing", "declined nursing",
                "refused leave", "declined leave",
                "refused activities", "declined activities"
            ],
            "Insight": [
                "insight", "awareness", "understanding", "recogni", "acknowledge", "accept",
                "denial", "denies", "lack of insight", "poor insight", "good insight",
                "partial insight", "limited insight", "believes", "doesn't believe",
                "mental illness", "unwell", "illness awareness"
            ],
            "Engagement": [
                "engag", "disengag", "attend", "did not attend", "dna", "missed appointment",
                "rapport", "therapeutic relationship", "working with", "cooperat",
                "uncooperative", "resistant", "reluctant", "willing", "unwilling",
                "participat", "involved", "motivation", "motivated"
            ],
            # Risk patterns from violence_risk_analyzer
            "Verbal Aggression": [
                "verbally abusive", "verbal abuse", "verbally aggressive", "verbal aggression",
                "shouting at", "shouted at", "raised voice", "screaming at",
                "swearing at", "swore at", "told to fuck off", "fuck off",
                "threatening language", "made threat", "making threat",
                "name calling", "name-calling", "called staff", "racial abuse", "racially abusive",
                "spat at", "spitting at", "spit at", "intimidat",
                "charged at", "squared up", "got in face"
            ],
            "Physical Aggression": [
                "physically aggressive", "physical aggression", "physically violent",
                "assaulted", "assault on", "attacked", "punched", "kicked", "hit staff",
                "slapped", "struck", "headbutt", "head butt", "bit staff", "bitten",
                "scratched staff", "grabbed staff", "pushed staff", "shoved",
                "threw at", "thrown at", "lashed out", "violent outburst", "violent episode",
                "restrained", "restraint required", "restraint needed", "prone restraint",
                "rapid tranquil", "rt administered", "prn given for aggression",
                "physical altercation", "physical intervention"
            ],
            "Property Damage": [
                "broke window", "broken window", "smashed", "damaged furniture",
                "punched wall", "kicked door", "threw furniture", "destroyed room",
                "vandal", "trashed room", "overturned table", "damage to property"
            ],
            "Self-Harm": [
                "self-harm", "self harm", "selfharm", "cut himself", "cut herself",
                "cutting", "laceration", "head banging", "banged head", "hit head against",
                "hitting self", "hitting himself", "hitting herself", "scratching self",
                "ligature", "attempted to hang", "overdose", "swallowed object",
                "threatened to harm himself", "threatened to harm herself",
                "threatened to kill himself", "threatened to kill herself",
                "suicidal ideation", "suicidal thoughts"
            ],
            "Sexual Behaviour": [
                "sexual comment", "inappropriate comment", "sexual remark",
                "sexual touch", "inappropriate touch", "grope", "groping",
                "exposed himself", "exposed herself", "exposure", "flashing",
                "masturbat", "naked in", "walking naked", "sexually disinhibit",
                "sexual advance", "inappropriate advance", "tried to kiss"
            ],
            "Bullying": [
                "bullying", "bully peer", "bullied", "targeting vulnerable",
                "took food from", "took cigarette from", "stealing from peer",
                "demanding money", "extort", "exploit", "intimidating peer",
                "coercing", "pressuring peer", "pushing peer", "shoving peer"
            ],
            "Self-Neglect": [
                "unkempt", "dishevelled", "unwashed", "dirty clothes", "soiled clothes",
                "body odour", "malodorous", "refused shower", "declined shower",
                "refused self-care", "poor self-care", "requires prompting",
                "poor room state", "room in mess", "not eating", "refused food",
                "weight loss", "dehydrat"
            ],
            "AWOL": [
                "awol", "absent without leave", "absconded", "absconding",
                "failed to return", "did not return from leave", "missing",
                "escaped", "left without permission", "breach of leave"
            ],
            "Substance Misuse": [
                "positive drug", "positive urine", "drug test positive",
                "cannabis", "cocaine", "heroin", "amphetamine", "spice",
                "alcohol on breath", "smelt of alcohol", "intoxicated",
                "suspected substance", "illicit substance", "found drugs"
            ],
        }

        # Category colors
        CATEGORY_COLORS = {
            "Compliance": "#059669",      # Green
            "Noncompliance": "#dc2626",   # Red
            "Admission": "#7c3aed",       # Purple
            "Insight": "#0891b2",         # Cyan
            "Engagement": "#2563eb",      # Blue
            "Verbal Aggression": "#e74c3c",    # Red-orange
            "Physical Aggression": "#b71c1c",  # Dark red
            "Property Damage": "#e53935",      # Red
            "Self-Harm": "#ff5722",            # Deep orange
            "Sexual Behaviour": "#e91e63",     # Pink
            "Bullying": "#795548",             # Brown
            "Self-Neglect": "#607d8b",         # Blue-grey
            "AWOL": "#f57c00",                 # Orange
            "Substance Misuse": "#9c27b0",     # Purple
        }

        # False positive exclusion phrases - avoid matching risk assessments, not actual incidents
        FALSE_POSITIVE_PHRASES = [
            "was not", "were not", "has not been", "had not been", "wasn't", "weren't",
            "no evidence of", "no signs of", "no indication of", "denied",
            "risk of", "at risk of", "risk assessment", "level of risk",
            "potential for", "possibility of", "likelihood of",
            "remains a risk", "continues to pose", "history of",
            "background of", "previous history", "known history",
            "there was no", "there were no", "nil", "none noted",
            "remained calm", "remained settled", "was calm", "was settled"
        ]

        def is_false_positive(text, keyword):
            """Check if the keyword match is a false positive based on context."""
            text_lower = text.lower()
            kw_lower = keyword.lower()

            # Find position of keyword
            pos = text_lower.find(kw_lower)
            if pos == -1:
                return False

            # Get context before the keyword (60 chars)
            context_start = max(0, pos - 60)
            context_before = text_lower[context_start:pos]

            # Check for false positive phrases in context before
            for fp_phrase in FALSE_POSITIVE_PHRASES:
                if fp_phrase in context_before:
                    return True

            return False

        # Get timeline episodes if available (for admission tagging)
        admission_periods = []
        if hasattr(self, '_timeline_episodes') and self._timeline_episodes:
            for ep in self._timeline_episodes:
                if ep.get("type") == "inpatient":
                    admission_periods.append((ep.get("start"), ep.get("end")))
            print(f"[MOJ-ASR] Section 4 popup: Found {len(admission_periods)} admission periods from timeline")

        def is_during_admission(entry_date):
            """Check if a date falls within any admission period."""
            if not entry_date or not admission_periods:
                return False
            # Convert to date object if datetime
            if isinstance(entry_date, datetime):
                check_date = entry_date.date()
            else:
                check_date = entry_date
            for start, end in admission_periods:
                if start and end and start <= check_date <= end:
                    return True
            return False

        # First pass: parse all dates and find the latest entry
        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})

            # Track latest date
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        # Use latest entry date as reference, look back 12 months
        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            # Fallback to now if no dates found
            cutoff_date = datetime.now() - timedelta(days=365)

        # Filter to 12-month window from latest entry
        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            # Include if date is within window or if no date (assume relevant)
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            self.popup_behaviour_import_layout.addWidget(placeholder)
            return

        print(f"[MOJ-ASR] Section 4 popup: {len(filtered_entries)} entries after 12-month filter")

        def extract_relevant_snippet(text, keywords_matched):
            """Extract first 2 lines for context + the sentence containing the keyword."""
            lines = text.strip().split('\n')
            # Get first 2 lines for context
            context_lines = lines[:2]
            context = '\n'.join(context_lines).strip()

            # Find sentences containing matched keywords
            relevant_sentences = []
            # Split into sentences
            import re
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for cat, keywords in CATEGORIES.items():
                if cat in keywords_matched:
                    for kw in keywords:
                        for sent in sentences:
                            if kw in sent.lower() and sent.strip() not in relevant_sentences:
                                # Don't duplicate if already in context
                                if sent.strip() not in context:
                                    relevant_sentences.append(sent.strip())
                                break

            # Build snippet
            snippet = context
            if relevant_sentences:
                # Add separator and relevant sentences
                additional = ' ... '.join(relevant_sentences[:2])  # Max 2 relevant sentences
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"

            return snippet if snippet else text[:200]

        # Categorize entries
        def categorize_text(text, entry_date):
            """Return list of matching categories for text and matched keywords info."""
            text_lower = text.lower()
            matches = []
            matched_keywords = {}

            # Check text-based categories with false positive filtering
            for cat, keywords in CATEGORIES.items():
                for kw in keywords:
                    if kw in text_lower:
                        # Check if this is a false positive (e.g., "no aggression", "risk of")
                        if not is_false_positive(text, kw):
                            matches.append(cat)
                            matched_keywords[cat] = kw
                            break

            # Check if during admission (date-based)
            if is_during_admission(entry_date) and "Admission" not in matches:
                matches.append("Admission")
                matched_keywords["Admission"] = "(during admission period)"

            return matches, matched_keywords

        categorized = []  # (entry, categories, matched_keywords)
        uncategorized = []

        for entry in filtered_entries:
            cats, matched_kws = categorize_text(entry["text"], entry.get("date_obj"))
            if cats:
                # Extract relevant snippet
                snippet = extract_relevant_snippet(entry["text"], cats)
                entry["snippet"] = snippet
                entry["matched_keywords"] = matched_kws
                categorized.append((entry, cats))
            else:
                uncategorized.append(entry)

        # Sort categorized by date (most recent first)
        categorized.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        # Sort uncategorized by date (most recent first)
        uncategorized.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)

        print(f"[MOJ-ASR] Section 4 popup: {len(categorized)} categorized, {len(uncategorized)} uncategorized")

        # Add category labels at top of import layout showing what categories were found (clickable to filter)
        # Remove old category labels container if exists
        if hasattr(self, '_behaviour_category_labels_container') and self._behaviour_category_labels_container:
            self._behaviour_category_labels_container.deleteLater()
            self._behaviour_category_labels_container = None

        # Store categorized data for filtering
        self._behaviour_all_categorized = categorized
        self._behaviour_categories = CATEGORIES
        self._behaviour_colors = CATEGORY_COLORS

        # Collect all unique categories found with counts
        category_counts = {}
        for entry, cats in categorized:
            for cat in cats:
                category_counts[cat] = category_counts.get(cat, 0) + 1

        if category_counts:
            # Create scrollable container for category labels
            labels_scroll = QScrollArea()
            labels_scroll.setWidgetResizable(True)
            labels_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            labels_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            labels_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            labels_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; } QScrollBar:horizontal { height: 6px; }")
            labels_scroll.setFixedHeight(36)

            labels_container = QWidget()
            labels_container.setStyleSheet("background: transparent;")
            labels_layout = QHBoxLayout(labels_container)
            labels_layout.setContentsMargins(0, 0, 0, 4)
            labels_layout.setSpacing(4)

            found_lbl = QLabel("Found:")
            found_lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #806000; background: transparent;")
            labels_layout.addWidget(found_lbl)

            for cat in sorted(category_counts.keys()):
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                count = category_counts[cat]
                btn = QPushButton(f"{cat} ({count})")
                btn.setCursor(Qt.CursorShape.PointingHandCursor)
                btn.setStyleSheet(f"QPushButton {{ font-size: 14px; font-weight: 600; color: white; background: {color}; padding: 2px 6px; border-radius: 4px; border: none; }} QPushButton:hover {{ opacity: 0.8; }}")
                btn.clicked.connect(lambda checked, c=cat: self._apply_behaviour_filter(c))
                labels_layout.addWidget(btn)

            labels_layout.addStretch()
            labels_scroll.setWidget(labels_container)
            self.popup_behaviour_import_layout.insertWidget(0, labels_scroll)
            self._behaviour_category_labels_container = labels_scroll

        def format_date_nice(date_obj):
            """Format date as '1st Dec 2025' style."""
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj  # Return as-is if string

            day = date_obj.day
            # Ordinal suffix
            if 11 <= day <= 13:
                suffix = "th"
            else:
                suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")

            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            """Convert hex color to a lighter background-friendly version."""
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            # Lighten by mixing with white (70% white, 30% original)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            """Highlight keywords in text that triggered the category matches."""
            import html
            escaped = html.escape(text)

            # Build keyword -> color mapping
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            # Sort by length (longest first) to avoid partial replacements
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)

            # Highlight each keyword (case-insensitive) with its category color
            import re
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )

            # Convert newlines to <br> for proper display
            escaped = escaped.replace('\n', '<br>')
            return escaped

        # Helper to create entry widget (4a house style)
        def create_entry_widget(entry, categories=None):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="attitude_behaviour",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                parent_container=getattr(self, 'popup_behaviour_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        # Only include categorized entries (those with labels) - sorted by date (most recent first)
        # Uncategorized entries are excluded from the output
        all_entries = [(entry, cats) for entry, cats in categorized]
        all_entries.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        if not all_entries:
            placeholder = QLabel("No categorized notes found (compliance, insight, engagement, or admissions).")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            self.popup_behaviour_import_layout.addWidget(placeholder)
            print(f"[MOJ-ASR] Section 4 popup: No categorized entries to display (filtered out {len(uncategorized)} unlabeled entries)")
            return

        # Add categorized entries in date order
        for entry, cats in all_entries:
            frame, cb, text, date = create_entry_widget(entry, cats)
            self.popup_behaviour_import_layout.addWidget(frame)
            self.popup_behaviour_imported_entries.append({"checkbox": cb, "text": text, "date": date, "categories": cats})

        print(f"[MOJ-ASR] Section 4 popup: Added {len(self.popup_behaviour_imported_entries)} labeled entries (excluded {len(uncategorized)} unlabeled)")

    def _apply_behaviour_filter(self, category):
        """Apply or remove category filter for behaviour imports (section 4)."""
        import re
        if not hasattr(self, '_behaviour_all_categorized') or not hasattr(self, 'popup_behaviour_import_layout'):
            return

        import_layout = self.popup_behaviour_import_layout

        # Clear existing entries (but keep the category labels at index 0)
        while import_layout.count() > 1:
            item = import_layout.takeAt(1)
            if item.widget():
                item.widget().deleteLater()

        self.popup_behaviour_imported_entries = []

        CATEGORIES = self._behaviour_categories
        CATEGORY_COLORS = self._behaviour_colors
        categorized = self._behaviour_all_categorized

        # Filter entries by category if specified
        if category:
            filtered = [(entry, cats) for entry, cats in categorized if category in cats]
        else:
            filtered = categorized

        if not filtered:
            placeholder = QLabel(f"No entries found for '{category}'." if category else "No categorized entries.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Add filter indicator with remove button if filtered (fixed height for consistency)
        if category:
            filter_frame = QFrame()
            filter_frame.setFixedHeight(28)
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 0)
            filter_row.setSpacing(8)

            color = CATEGORY_COLORS.get(category, "#6b7280")
            filter_label = QLabel(f"Filtered by: {category}")
            filter_label.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(filter_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setFixedHeight(22)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 4px; padding: 2px 8px;
                }
                QPushButton:hover { background: #fef2f2; }
            """)
            remove_btn.clicked.connect(lambda: self._apply_behaviour_filter(None))
            filter_row.addWidget(remove_btn)

            import_layout.addWidget(filter_frame)

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>', escaped)
            escaped = escaped.replace('\n', '<br>')
            return escaped

        for entry, cats in filtered:
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, cats or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="attitude_behaviour",
                categories=cats,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                parent_container=getattr(self, 'popup_behaviour_import_container', None),
            )

            import_layout.addWidget(entry_frame)
            self.popup_behaviour_imported_entries.append({"checkbox": cb, "text": full_text, "date": date_raw, "categories": cats})

    def populate_popup_patient_attitude_imports(self, entries: list):
        """Populate the imported data panel in popup 6 (Patient's Attitude) with categorized entries.

        Categories: Index Offence, Offending Behaviour, Remorse, Victim Empathy, Compliance.
        Filters to 12-month window from latest entry. Only shows labeled entries with findings.
        Clickable labels filter by category.
        """
        from datetime import datetime, timedelta
        import re

        print(f"[MOJ-ASR] Section 6 popup: populate_popup_patient_attitude_imports called with {len(entries) if entries else 0} entries")

        # Check if UI elements exist
        import_layout = getattr(self, 'popup_patient_attitude_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 6 popup: popup_patient_attitude_import_layout not available")
            return

        # Clear existing entries
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_patient_attitude_imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            print("[MOJ-ASR] Section 6 popup: No entries to display")
            return

        # Category keywords for patient attitude (offending behaviour focus)
        CATEGORIES = {
            "Index Offence": [
                "index offence", "index offense", "original offence", "original offense",
                "manslaughter", "murder", "attempted murder", "gbh", "abh",
                "threats to kill", "wounding", "arson", "rape", "sexual assault",
                "sexual offence", "indecent assault"
            ],
            "Offending Behaviour": [
                "offence", "offense", "conviction", "prison", "court", "police"
            ],
            "Remorse": [
                "remorse", "remorseful", "regret", "sorry", "apologise", "apologize",
                "apology", "guilt", "guilty feelings", "ashamed", "shame", "sorrow",
                "sorry for", "feels bad", "expressed regret", "no remorse", "lack of remorse",
                "doesn't show remorse", "shows no remorse", "callous", "indifferent"
            ],
            "Victim Empathy": [
                "victim empathy", "empathy", "empathic", "victim awareness", "victim impact",
                "understands impact", "effect on victim", "harm caused", "hurt caused",
                "victim's perspective", "put themselves in", "sympathy", "compassion",
                "lack of empathy", "no empathy", "doesn't understand impact",
                "victim work", "victim letter", "restorative", "callous disregard"
            ],
            "Compliance": [
                "complian", "non-complian", "noncomplian", "adherent", "non-adherent",
                "medication", "meds", "depot", "clozapine", "refused", "refusing",
                "taking medication", "not taking", "stopped taking", "concordan",
                "engagement with treatment", "treatment adherence", "cooperat",
                "uncooperative", "resistant", "reluctant", "willing", "participation"
            ],
        }

        # Category colors
        CATEGORY_COLORS = {
            "Index Offence": "#7c3aed",      # Purple
            "Offending Behaviour": "#dc2626", # Red
            "Remorse": "#0891b2",             # Cyan
            "Victim Empathy": "#059669",      # Green
            "Compliance": "#d97706",          # Amber
        }

        # Store for filtering
        self._patient_attitude_categories = CATEGORIES
        self._patient_attitude_colors = CATEGORY_COLORS
        self._patient_attitude_current_filter = None

        # Parse all dates and find the latest entry
        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})

            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        # Use latest entry date as reference, look back 12 months
        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        # Filter to 12-month window
        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        print(f"[MOJ-ASR] Section 6 popup: {len(filtered_entries)} entries after 12-month filter")

        def extract_relevant_snippet(text, matched_categories):
            """Extract first 2 lines for context + sentence containing the keyword."""
            lines = text.strip().split('\n')
            context_lines = lines[:2]
            context = '\n'.join(context_lines).strip()

            relevant_sentences = []
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for cat in matched_categories:
                keywords = CATEGORIES.get(cat, [])
                for kw in keywords:
                    for sent in sentences:
                        if kw in sent.lower() and sent.strip() not in relevant_sentences:
                            if sent.strip() not in context:
                                relevant_sentences.append(sent.strip())
                            break

            snippet = context
            if relevant_sentences:
                additional = ' ... '.join(relevant_sentences[:2])
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"

            return snippet if snippet else text[:200]

        def categorize_text(text):
            """Return list of matching categories."""
            text_lower = text.lower()
            matches = []

            for cat, keywords in CATEGORIES.items():
                for kw in keywords:
                    if kw in text_lower:
                        matches.append(cat)
                        break

            return matches

        categorized = []
        for entry in filtered_entries:
            cats = categorize_text(entry["text"])
            if cats:
                snippet = extract_relevant_snippet(entry["text"], cats)
                entry["snippet"] = snippet
                categorized.append((entry, cats))

        categorized.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        # Add category labels to header showing what categories were found (clickable to filter)
        category_buttons_layout = getattr(self, 'popup_patient_attitude_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            category_counts = {}
            for entry, cats in categorized:
                for cat in cats:
                    category_counts[cat] = category_counts.get(cat, 0) + 1
            if category_counts:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
            for cat in sorted(category_counts.keys()):
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                count = category_counts[cat]
                btn = QPushButton(f"{cat} ({count})")
                btn.setCursor(Qt.CursorShape.PointingHandCursor)
                btn.setStyleSheet(f"QPushButton {{ font-size: 14px; font-weight: 600; color: white; background: {color}; padding: 2px 6px; border-radius: 4px; border: none; }} QPushButton:hover {{ opacity: 0.8; }}")
                btn.clicked.connect(lambda checked, c=cat: self._apply_patient_attitude_filter(c))
                category_buttons_layout.addWidget(btn)

        if not categorized:
            placeholder = QLabel("No relevant notes found (index offence, offending behaviour, remorse, empathy, compliance).")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            print(f"[MOJ-ASR] Section 6 popup: No categorized entries to display")
            return

        # Store all categorized entries for filtering
        self._patient_attitude_all_categorized = categorized

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            if 11 <= day <= 13:
                suffix = "th"
            else:
                suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            """Convert hex color to a lighter background-friendly version."""
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            # Lighten by mixing with white (70% white, 30% original)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            """Highlight keywords in text that triggered the category matches."""
            import html
            escaped = html.escape(text)

            # Build keyword -> color mapping
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            # Sort by length (longest first) to avoid partial replacements
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)

            # Highlight each keyword (case-insensitive) with its category color
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )

            # Convert newlines to <br> for proper display
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="patient_attitude",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_patient_attitude_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        # Display entries with optional filter
        self._display_patient_attitude_entries(import_layout, categorized, create_entry_widget, format_date_nice, None)

        print(f"[MOJ-ASR] Section 6 popup: Added {len(self.popup_patient_attitude_imported_entries)} labeled entries")

    def _display_patient_attitude_entries(self, import_layout, categorized, create_entry_widget, format_date_nice, filter_category):
        """Display patient attitude entries with optional category filter."""
        # Clear existing entries
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_patient_attitude_imported_entries = []

        # Filter header row (only show when filter is active - fixed height for consistency)
        if filter_category:
            filter_frame = QFrame()
            filter_frame.setFixedHeight(28)
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 0)
            filter_row.setSpacing(8)

            color = self._patient_attitude_colors.get(filter_category, "#6b7280")
            filter_label = QLabel(f"Filtered by: {filter_category}")
            filter_label.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(filter_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setFixedHeight(22)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 4px; padding: 2px 8px;
                }
                QPushButton:hover { background: #fef2f2; }
            """)
            remove_btn.clicked.connect(lambda: self._apply_patient_attitude_filter(None))
            filter_row.addWidget(remove_btn)

            import_layout.addWidget(filter_frame)

        # Filter entries if needed
        if filter_category:
            filtered = [(entry, cats) for entry, cats in categorized if filter_category in cats]
        else:
            filtered = categorized

        if not filtered:
            placeholder = QLabel(f"No entries found for '{filter_category}'." if filter_category else "No entries found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Create entry widgets
        for entry, cats in filtered:
            frame, cb, text, date = create_entry_widget(entry, cats, self._apply_patient_attitude_filter)
            import_layout.addWidget(frame)
            self.popup_patient_attitude_imported_entries.append({"checkbox": cb, "text": text, "date": date, "categories": cats, "frame": frame})

    def _apply_patient_attitude_filter(self, category):
        """Apply or remove category filter for patient attitude imports."""
        import re
        import_layout = getattr(self, 'popup_patient_attitude_import_layout', None)
        if not import_layout or not hasattr(self, '_patient_attitude_all_categorized'):
            return

        self._patient_attitude_current_filter = category

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            if 11 <= day <= 13:
                suffix = "th"
            else:
                suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        CATEGORY_COLORS = self._patient_attitude_colors
        CATEGORIES = self._patient_attitude_categories

        def hex_to_light_bg(hex_color):
            """Convert hex color to a lighter background-friendly version."""
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            # Lighten by mixing with white (70% white, 30% original)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            """Highlight keywords in text that triggered the category matches."""
            import html
            escaped = html.escape(text)

            # Build keyword -> color mapping
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            # Sort by length (longest first) to avoid partial replacements
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)

            # Highlight each keyword (case-insensitive) with its category color
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )

            # Convert newlines to <br> for proper display
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="patient_attitude",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_patient_attitude_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_patient_attitude_entries(
            import_layout,
            self._patient_attitude_all_categorized,
            create_entry_widget,
            format_date_nice,
            category
        )

    def populate_popup_capacity_imports(self, entries: list):
        """Populate the imported data panel in popup 7 (Capacity) with categorized entries.

        Categories: Capacity, Best Interest, IMCA, DoLS, COP, SOAD, Appointeeship, Guardianship, IMHA.
        Filters to 12-month window from latest entry. Only shows labeled entries with findings.
        Clickable labels filter by category.
        """
        from datetime import datetime, timedelta
        import re

        print(f"[MOJ-ASR] Section 7 popup: populate_popup_capacity_imports called with {len(entries) if entries else 0} entries")

        import_layout = getattr(self, 'popup_capacity_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 7 popup: popup_capacity_import_layout not available")
            return

        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_capacity_imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Category keywords for capacity
        CATEGORIES = {
            "Capacity": [
                "capacity", "lacks capacity", "has capacity", "mental capacity",
                "capacity assessment", "capacity to consent", "incapacitated",
                "decision making", "unable to make decisions"
            ],
            "Best Interest": [
                "best interest", "best interests", "bi meeting", "bi decision"
            ],
            "IMCA": [
                "imca", "independent mental capacity advocate"
            ],
            "DoLS": [
                "dols", "deprivation of liberty", "liberty protection safeguard", "lps"
            ],
            "COP": [
                "court of protection", "cop application", "cop order"
            ],
            "SOAD": [
                "soad", "second opinion", "second opinion appointed doctor"
            ],
            "Appointeeship": [
                "appointee", "appointeeship", "dwp appointee"
            ],
            "Guardianship": [
                "guardianship", "guardian"
            ],
            "IMHA": [
                "imha", "independent mental health advocate"
            ],
            "Finances": [
                "finances", "financial", "money", "benefits", "pension", "pip", "esa"
            ],
            "Residence": [
                "residence", "accommodation", "placement", "discharge destination"
            ]
        }

        CATEGORY_COLORS = {
            "Capacity": "#7c3aed",
            "Best Interest": "#dc2626",
            "IMCA": "#0891b2",
            "DoLS": "#059669",
            "COP": "#d97706",
            "SOAD": "#be185d",
            "Appointeeship": "#4f46e5",
            "Guardianship": "#0d9488",
            "IMHA": "#ea580c",
            "Finances": "#6366f1",
            "Residence": "#84cc16"
        }

        self._capacity_categories = CATEGORIES
        self._capacity_colors = CATEGORY_COLORS
        self._capacity_current_filter = None

        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def extract_relevant_snippet(text, matched_categories):
            lines = text.strip().split('\n')
            context = '\n'.join(lines[:2]).strip()
            relevant_sentences = []
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for cat in matched_categories:
                keywords = CATEGORIES.get(cat, [])
                for kw in keywords:
                    for sent in sentences:
                        if kw in sent.lower() and sent.strip() not in relevant_sentences:
                            if sent.strip() not in context:
                                relevant_sentences.append(sent.strip())
                            break

            snippet = context
            if relevant_sentences:
                additional = ' ... '.join(relevant_sentences[:2])
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"
            return snippet if snippet else text[:200]

        def categorize_text(text):
            text_lower = text.lower()
            matches = []
            for cat, keywords in CATEGORIES.items():
                for kw in keywords:
                    if kw in text_lower:
                        matches.append(cat)
                        break
            return matches

        categorized = []
        for entry in filtered_entries:
            cats = categorize_text(entry["text"])
            if cats:
                snippet = extract_relevant_snippet(entry["text"], cats)
                entry["snippet"] = snippet
                categorized.append((entry, cats))

        categorized.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        # Add category labels to header showing what categories were found (clickable to filter)
        category_buttons_layout = getattr(self, 'popup_capacity_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            category_counts = {}
            for entry, cats in categorized:
                for cat in cats:
                    category_counts[cat] = category_counts.get(cat, 0) + 1
            if category_counts:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
            for cat in sorted(category_counts.keys()):
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                count = category_counts[cat]
                btn = QPushButton(f"{cat} ({count})")
                btn.setCursor(Qt.CursorShape.PointingHandCursor)
                btn.setStyleSheet(f"QPushButton {{ font-size: 14px; font-weight: 600; color: white; background: {color}; padding: 2px 6px; border-radius: 4px; border: none; }} QPushButton:hover {{ opacity: 0.8; }}")
                btn.clicked.connect(lambda checked, c=cat: self._apply_capacity_filter(c))
                category_buttons_layout.addWidget(btn)

        if not categorized:
            placeholder = QLabel("No capacity-related notes found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        self._capacity_all_categorized = categorized

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            """Convert hex color to a lighter background-friendly version."""
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            # Lighten by mixing with white (70% white, 30% original)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            """Highlight keywords in text that triggered the category matches."""
            import html
            escaped = html.escape(text)

            # Build keyword -> color mapping
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            # Sort by length (longest first) to avoid partial replacements
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)

            # Highlight each keyword (case-insensitive) with its category color
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )

            # Convert newlines to <br> for proper display
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="capacity",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_capacity_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_capacity_entries(import_layout, categorized, create_entry_widget, format_date_nice, None)
        print(f"[MOJ-ASR] Section 7 popup: Added {len(self.popup_capacity_imported_entries)} labeled entries")

    def _display_capacity_entries(self, import_layout, categorized, create_entry_widget, format_date_nice, filter_category):
        """Display capacity entries with optional category filter."""
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_capacity_imported_entries = []

        # Fixed height filter bar for consistency
        if filter_category:
            filter_frame = QFrame()
            filter_frame.setFixedHeight(28)
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 0)
            filter_row.setSpacing(8)

            color = self._capacity_colors.get(filter_category, "#6b7280")
            filter_label = QLabel(f"Filtered by: {filter_category}")
            filter_label.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(filter_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setFixedHeight(22)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 4px; padding: 2px 8px;
                }
                QPushButton:hover { background: #fef2f2; }
            """)
            remove_btn.clicked.connect(lambda: self._apply_capacity_filter(None))
            filter_row.addWidget(remove_btn)

            import_layout.addWidget(filter_frame)

        if filter_category:
            filtered = [(entry, cats) for entry, cats in categorized if filter_category in cats]
        else:
            filtered = categorized

        if not filtered:
            placeholder = QLabel(f"No entries found for '{filter_category}'." if filter_category else "No entries found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        for entry, cats in filtered:
            frame, cb, text, date = create_entry_widget(entry, cats, self._apply_capacity_filter)
            import_layout.addWidget(frame)
            self.popup_capacity_imported_entries.append({"checkbox": cb, "text": text, "date": date, "categories": cats, "frame": frame})

    def _apply_capacity_filter(self, category):
        """Apply or remove category filter for capacity imports."""
        import re
        import_layout = getattr(self, 'popup_capacity_import_layout', None)
        if not import_layout or not hasattr(self, '_capacity_all_categorized'):
            return

        self._capacity_current_filter = category
        CATEGORY_COLORS = self._capacity_colors
        CATEGORIES = self._capacity_categories

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            """Convert hex color to a lighter background-friendly version."""
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            # Lighten by mixing with white (70% white, 30% original)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            """Highlight keywords in text that triggered the category matches."""
            import html
            escaped = html.escape(text)

            # Build keyword -> color mapping
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            # Sort by length (longest first) to avoid partial replacements
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)

            # Highlight each keyword (case-insensitive) with its category color
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )

            # Convert newlines to <br> for proper display
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="capacity",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_capacity_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_capacity_entries(
            import_layout,
            self._capacity_all_categorized,
            create_entry_widget,
            format_date_nice,
            category
        )

    def populate_popup_risk_addressed_imports(self, entries: list):
        """Populate the imported data panel in popup 10 (How Risks Addressed) with risk factor entries.

        Categories based on Section 9 risk types: Violence, Self-harm, Suicide, Self-neglect,
        Exploitation, Substance misuse, Deterioration, Non-compliance, Absconding, Reoffending.
        Filters to 12-month window from latest entry.
        """
        from datetime import datetime, timedelta
        import re

        print(f"[MOJ-ASR] Section 10 popup: populate_popup_risk_addressed_imports called with {len(entries) if entries else 0} entries")

        import_layout = getattr(self, 'popup_risk_addressed_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 10 popup: popup_risk_addressed_import_layout not available")
            return

        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_risk_addressed_imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Category keywords matching Section 9 risk types
        CATEGORIES = {
            "Violence": [
                "violen", "aggress", "assault", "attack", "threat", "intimidat",
                "hostil", "physical altercation", "hit", "punch", "weapon",
                "danger to others", "harm to others"
            ],
            "Self-harm": [
                "self-harm", "self harm", "cutting", "self-injur", "self injur",
                "overdose", "ligature", "scratch", "burn", "wound"
            ],
            "Suicide": [
                "suicid", "end my life", "kill myself", "want to die", "death wish",
                "suicidal ideation", "thoughts of death", "plan to end"
            ],
            "Self-neglect": [
                "self-neglect", "self neglect", "poor hygiene", "not eating",
                "weight loss", "dehydrat", "malnutrition", "refusing food",
                "not caring for", "unkempt", "dirty", "neglecting"
            ],
            "Exploitation": [
                "exploit", "vulnerab", "taken advantage", "financial abuse",
                "manipulat", "cuckooing", "coerced", "grooming", "used by others"
            ],
            "Substance": [
                "substance", "drug", "alcohol", "cannabis", "cocaine", "heroin",
                "spice", "intoxicat", "under the influence", "illicit", "misuse",
                "addiction", "drinking", "using drugs", "relapse"
            ],
            "Deterioration": [
                "deteriorat", "relapse", "worsening", "decline", "decompensate",
                "mental state decline", "becoming unwell", "acutely unwell",
                "symptoms returning", "psychotic", "paranoi"
            ],
            "Non-compliance": [
                "non-complia", "noncomplian", "refused medication", "not taking",
                "declined treatment", "disengag", "not engaging", "missed",
                "did not attend", "dna", "non-adherent", "nonadherent"
            ],
            "Absconding": [
                "abscond", "awol", "absent without leave", "missing", "escaped",
                "left without permission", "failed to return", "unauthorised absence"
            ],
            "Reoffending": [
                "reoffend", "re-offend", "further offence", "new offence", "arrested",
                "charged", "conviction", "criminal", "police involvement", "court"
            ],
        }

        CATEGORY_COLORS = {
            "Violence": "#dc2626",       # Red
            "Self-harm": "#ea580c",      # Orange
            "Suicide": "#7c3aed",        # Purple
            "Self-neglect": "#0891b2",   # Cyan
            "Exploitation": "#be185d",   # Pink
            "Substance": "#059669",      # Green
            "Deterioration": "#d97706",  # Amber
            "Non-compliance": "#4f46e5", # Indigo
            "Absconding": "#0d9488",     # Teal
            "Reoffending": "#991b1b",    # Dark red
        }

        self._risk_addressed_categories = CATEGORIES
        self._risk_addressed_colors = CATEGORY_COLORS
        self._risk_addressed_current_filter = None

        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def extract_relevant_snippet(text, matched_categories):
            lines = text.strip().split('\n')
            context = '\n'.join(lines[:2]).strip()
            relevant_sentences = []
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for cat in matched_categories:
                keywords = CATEGORIES.get(cat, [])
                for kw in keywords:
                    for sent in sentences:
                        if kw in sent.lower() and sent.strip() not in relevant_sentences:
                            if sent.strip() not in context:
                                relevant_sentences.append(sent.strip())
                            break

            snippet = context
            if relevant_sentences:
                additional = ' ... '.join(relevant_sentences[:2])
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"
            return snippet if snippet else text[:200]

        def categorize_text(text):
            text_lower = text.lower()
            matches = []
            # Keywords that need whole-word matching to avoid false positives
            WHOLE_WORD_KEYWORDS = {"charged"}
            for cat, keywords in CATEGORIES.items():
                for kw in keywords:
                    if kw in WHOLE_WORD_KEYWORDS:
                        # Use word boundary regex
                        if re.search(r'\b' + re.escape(kw) + r'\b', text_lower):
                            matches.append(cat)
                            break
                    else:
                        if kw in text_lower:
                            matches.append(cat)
                            break
            return matches

        categorized = []
        for entry in filtered_entries:
            cats = categorize_text(entry["text"])
            if cats:
                snippet = extract_relevant_snippet(entry["text"], cats)
                entry["snippet"] = snippet
                categorized.append((entry, cats))

        categorized.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        # Add category labels to header showing what categories were found (clickable to filter)
        category_buttons_layout = getattr(self, 'popup_risk_addressed_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            category_counts = {}
            for entry, cats in categorized:
                for cat in cats:
                    category_counts[cat] = category_counts.get(cat, 0) + 1
            if category_counts:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
            for cat in sorted(category_counts.keys()):
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                count = category_counts[cat]
                btn = QPushButton(f"{cat} ({count})")
                btn.setCursor(Qt.CursorShape.PointingHandCursor)
                btn.setStyleSheet(f"QPushButton {{ font-size: 14px; font-weight: 600; color: white; background: {color}; padding: 2px 6px; border-radius: 4px; border: none; }} QPushButton:hover {{ opacity: 0.8; }}")
                btn.clicked.connect(lambda checked, c=cat: self._apply_risk_addressed_filter(c))
                category_buttons_layout.addWidget(btn)

        if not categorized:
            placeholder = QLabel("No risk-related notes found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        self._risk_addressed_all_categorized = categorized

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="risk_addressed",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_risk_addressed_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_risk_addressed_entries(import_layout, categorized, create_entry_widget, format_date_nice, None)
        print(f"[MOJ-ASR] Section 10 popup: Added {len(self.popup_risk_addressed_imported_entries)} labeled entries")

    def _display_risk_addressed_entries(self, import_layout, categorized, create_entry_widget, format_date_nice, filter_category):
        """Display risk addressed entries with optional category filter."""
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_risk_addressed_imported_entries = []

        # Fixed height filter bar for consistency
        if filter_category:
            filter_frame = QFrame()
            filter_frame.setFixedHeight(28)
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 0)
            filter_row.setSpacing(8)

            color = self._risk_addressed_colors.get(filter_category, "#6b7280")
            filter_label = QLabel(f"Filtered by: {filter_category}")
            filter_label.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(filter_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setFixedHeight(22)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 4px; padding: 2px 8px;
                }
                QPushButton:hover { background: #fef2f2; }
            """)
            remove_btn.clicked.connect(lambda: self._apply_risk_addressed_filter(None))
            filter_row.addWidget(remove_btn)

            import_layout.addWidget(filter_frame)

        if filter_category:
            filtered = [(entry, cats) for entry, cats in categorized if filter_category in cats]
        else:
            filtered = categorized

        if not filtered:
            placeholder = QLabel(f"No entries found for '{filter_category}'." if filter_category else "No entries found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        for entry, cats in filtered:
            frame, cb, text, date = create_entry_widget(entry, cats, self._apply_risk_addressed_filter)
            import_layout.addWidget(frame)
            self.popup_risk_addressed_imported_entries.append({"checkbox": cb, "text": text, "date": date, "categories": cats, "frame": frame})

    def _apply_risk_addressed_filter(self, category):
        """Apply or remove category filter for risk addressed imports."""
        import re
        import_layout = getattr(self, 'popup_risk_addressed_import_layout', None)
        if not import_layout or not hasattr(self, '_risk_addressed_all_categorized'):
            return

        self._risk_addressed_current_filter = category
        CATEGORY_COLORS = self._risk_addressed_colors
        CATEGORIES = self._risk_addressed_categories

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="risk_addressed",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_risk_addressed_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_risk_addressed_entries(
            import_layout,
            self._risk_addressed_all_categorized,
            create_entry_widget,
            format_date_nice,
            category
        )

    # ============================================================
    # SECTION 11 - Abscond/Escape Imported Notes
    # ============================================================

    def populate_popup_abscond_imports(self, entries: list):
        """Populate the imported data panel in popup 11 (Abscond/Escape) with AWOL-related entries.

        Categories: AWOL, Absconding, Escape, Failure to Return
        Filters to 12-month window from latest entry.
        """
        from datetime import datetime, timedelta
        import re

        print(f"[MOJ-ASR] Section 11 popup: populate_popup_abscond_imports called with {len(entries) if entries else 0} entries")

        import_layout = getattr(self, 'popup_abscond_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 11 popup: popup_abscond_import_layout not available")
            return

        # Clear existing entries (skip first item which is the filter row)
        while import_layout.count() > 1:
            item = import_layout.takeAt(1)
            if item.widget():
                item.widget().deleteLater()

        self.popup_abscond_imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Category keywords for AWOL/absconding/escape
        CATEGORIES = {
            "AWOL": [
                "awol", "a.w.o.l", "absent without leave", "absent without official leave",
                "unauthorised absence", "unauthorized absence", "left without permission"
            ],
            "Absconding": [
                "abscond", "absconded", "absconding", "absconded from", "went missing",
                "whereabouts unknown", "failed to return to ward"
            ],
            "Escape": [
                "escape", "escaped", "escaping", "fled", "ran away", "ran off",
                "broke out", "got away", "evaded", "eluded"
            ],
            "Failure to Return": [
                "failed to return", "failure to return", "did not return", "didn't return",
                "overdue from leave", "not returned", "late returning", "late from leave",
                "breach of leave", "leave conditions"
            ],
        }

        CATEGORY_COLORS = {
            "AWOL": "#dc2626",            # Red
            "Absconding": "#0d9488",      # Teal
            "Escape": "#7c3aed",          # Purple
            "Failure to Return": "#d97706",  # Amber
        }

        self._abscond_categories = CATEGORIES
        self._abscond_colors = CATEGORY_COLORS
        self._abscond_current_filter = None

        # Update filter dropdown
        if hasattr(self, 'popup_abscond_filter'):
            self.popup_abscond_filter.blockSignals(True)
            self.popup_abscond_filter.clear()
            self.popup_abscond_filter.addItem("All")
            for cat in CATEGORIES.keys():
                self.popup_abscond_filter.addItem(cat)
            self.popup_abscond_filter.blockSignals(False)

        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def extract_relevant_snippet(text, matched_categories):
            lines = text.strip().split('\n')
            context = '\n'.join(lines[:2]).strip()
            relevant_sentences = []
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for cat in matched_categories:
                keywords = CATEGORIES.get(cat, [])
                for kw in keywords:
                    for sent in sentences:
                        if kw in sent.lower() and sent.strip() not in relevant_sentences:
                            if sent.strip() not in context:
                                relevant_sentences.append(sent.strip())
                            break

            snippet = context
            if relevant_sentences:
                additional = ' ... '.join(relevant_sentences[:2])
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"
            return snippet if snippet else text[:200]

        def categorize_text(text):
            text_lower = text.lower()
            matches = []
            for cat, keywords in CATEGORIES.items():
                for kw in keywords:
                    if kw in text_lower:
                        matches.append(cat)
                        break
            return matches

        categorized = []
        for entry in filtered_entries:
            cats = categorize_text(entry["text"])
            if cats:
                snippet = extract_relevant_snippet(entry["text"], cats)
                entry["snippet"] = snippet
                categorized.append((entry, cats))

        categorized.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        # Add category labels to header showing what categories were found (clickable to filter)
        category_buttons_layout = getattr(self, 'popup_abscond_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            category_counts = {}
            for entry, cats in categorized:
                for cat in cats:
                    category_counts[cat] = category_counts.get(cat, 0) + 1
            if category_counts:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
            for cat in sorted(category_counts.keys()):
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                count = category_counts[cat]
                btn = QPushButton(f"{cat} ({count})")
                btn.setCursor(Qt.CursorShape.PointingHandCursor)
                btn.setStyleSheet(f"QPushButton {{ font-size: 14px; font-weight: 600; color: white; background: {color}; padding: 2px 6px; border-radius: 4px; border: none; }} QPushButton:hover {{ opacity: 0.8; }}")
                btn.clicked.connect(lambda checked, c=cat: self._apply_abscond_filter_by_tag(c))
                category_buttons_layout.addWidget(btn)

        if not categorized:
            placeholder = QLabel("No AWOL/absconding/escape-related notes found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        self._abscond_all_categorized = categorized

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="abscond",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_abscond_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_abscond_entries(import_layout, categorized, create_entry_widget, format_date_nice, None)
        print(f"[MOJ-ASR] Section 11 popup: Added {len(self.popup_abscond_imported_entries)} labeled entries")

    def _display_abscond_entries(self, import_layout, categorized, create_entry_widget, format_date_nice, filter_category):
        """Display abscond entries with optional category filter."""
        # Clear existing entries (skip first item which is the filter row)
        while import_layout.count() > 1:
            item = import_layout.takeAt(1)
            if item.widget():
                item.widget().deleteLater()

        self.popup_abscond_imported_entries = []

        # Fixed height filter bar for consistency
        if filter_category:
            filter_frame = QFrame()
            filter_frame.setFixedHeight(28)
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 0)
            filter_row.setSpacing(8)

            color = self._abscond_colors.get(filter_category, "#6b7280")
            filter_label = QLabel(f"Filtered by: {filter_category}")
            filter_label.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(filter_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setFixedHeight(22)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 4px; padding: 2px 8px;
                }
                QPushButton:hover { background: #fef2f2; }
            """)
            remove_btn.clicked.connect(lambda: self._apply_abscond_filter_by_tag(None))
            filter_row.addWidget(remove_btn)

            import_layout.addWidget(filter_frame)

        if filter_category:
            filtered = [(entry, cats) for entry, cats in categorized if filter_category in cats]
        else:
            filtered = categorized

        if not filtered:
            placeholder = QLabel(f"No entries found for '{filter_category}'." if filter_category else "No entries found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        for entry, cats in filtered:
            frame, cb, text, date = create_entry_widget(entry, cats, self._apply_abscond_filter_by_tag)
            import_layout.addWidget(frame)
            self.popup_abscond_imported_entries.append({"checkbox": cb, "text": text, "date": date, "categories": cats, "frame": frame})

    def _apply_abscond_filter_by_tag(self, category):
        """Apply or remove category filter for abscond imports (triggered by tag click)."""
        import re
        import_layout = getattr(self, 'popup_abscond_import_layout', None)
        if not import_layout or not hasattr(self, '_abscond_all_categorized'):
            return

        self._abscond_current_filter = category
        CATEGORY_COLORS = self._abscond_colors
        CATEGORIES = self._abscond_categories

        # Update dropdown to match
        if hasattr(self, 'popup_abscond_filter'):
            self.popup_abscond_filter.blockSignals(True)
            if category:
                idx = self.popup_abscond_filter.findText(category)
                if idx >= 0:
                    self.popup_abscond_filter.setCurrentIndex(idx)
            else:
                self.popup_abscond_filter.setCurrentIndex(0)  # "All"
            self.popup_abscond_filter.blockSignals(False)

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color

            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="abscond",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_abscond_import_container', None),
            )
            return entry_frame, cb, full_text, date_raw

        self._display_abscond_entries(
            import_layout,
            self._abscond_all_categorized,
            create_entry_widget,
            format_date_nice,
            category
        )

    def _apply_abscond_filter(self, category_text):
        """Apply filter from dropdown selection."""
        if category_text == "All":
            self._apply_abscond_filter_by_tag(None)
        else:
            self._apply_abscond_filter_by_tag(category_text)

    # ============================================================
    # SECTION 12 - MAPPA Imported Notes
    # ============================================================

    def populate_popup_mappa_imports(self, entries: list):
        """Populate the imported data panel in popup 12 (MAPPA) with MAPPA-related entries.

        Searches for MAPPA references in clinical notes.
        Filters to 12-month window from latest entry.
        """
        from datetime import datetime, timedelta
        import re

        print(f"[MOJ-ASR] Section 12 popup: populate_popup_mappa_imports called with {len(entries) if entries else 0} entries")

        import_layout = getattr(self, 'popup_mappa_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 12 popup: popup_mappa_import_layout not available")
            return

        # Clear existing entries
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_mappa_imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # MAPPA-related keywords
        MAPPA_KEYWORDS = [
            "mappa", "multi-agency public protection", "multi agency public protection",
            "public protection", "mappa meeting", "mappa level", "mappa category",
            "mappa coordinator", "mappa panel", "mappa review", "visor",
            "sex offender", "violent offender", "public protection arrangement",
            "risk management meeting", "level 1", "level 2", "level 3",
            "category 1", "category 2", "category 3", "offender manager",
            "probation", "nps", "police public protection"
        ]

        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def extract_relevant_snippet(text, matched_keywords):
            lines = text.strip().split('\n')
            context = '\n'.join(lines[:2]).strip()
            relevant_sentences = []
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for kw in matched_keywords:
                for sent in sentences:
                    if kw in sent.lower() and sent.strip() not in relevant_sentences:
                        if sent.strip() not in context:
                            relevant_sentences.append(sent.strip())
                        break

            snippet = context
            if relevant_sentences:
                additional = ' ... '.join(relevant_sentences[:2])
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"
            return snippet if snippet else text[:200]

        def find_mappa_keywords(text):
            text_lower = text.lower()
            matches = []
            # Keywords that need whole-word matching to avoid false positives
            WHOLE_WORD_KEYWORDS = {"visor", "mappa", "nps"}
            for kw in MAPPA_KEYWORDS:
                if kw in WHOLE_WORD_KEYWORDS:
                    # Use word boundary regex for these keywords
                    if re.search(r'\b' + re.escape(kw) + r'\b', text_lower):
                        matches.append(kw)
                else:
                    # Standard substring match for phrases
                    if kw in text_lower:
                        matches.append(kw)
            return matches

        categorized = []
        for entry in filtered_entries:
            matched = find_mappa_keywords(entry["text"])
            if matched:
                snippet = extract_relevant_snippet(entry["text"], matched)
                entry["snippet"] = snippet
                entry["matched_keywords"] = matched
                categorized.append(entry)

        categorized.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)

        # Add category label to header showing count of MAPPA entries found
        category_buttons_layout = getattr(self, 'popup_mappa_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            if categorized:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
                lbl = QLabel(f"MAPPA ({len(categorized)})")
                lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: white; background: #7c3aed; padding: 2px 6px; border-radius: 4px;")
                category_buttons_layout.addWidget(lbl)

        if not categorized:
            placeholder = QLabel("No MAPPA-related notes found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_keywords):
            import html
            escaped = html.escape(text)
            mappa_color = "#7c3aed"  # Purple for MAPPA
            light_color = hex_to_light_bg(mappa_color)

            sorted_keywords = sorted(matched_keywords, key=len, reverse=True)
            for kw in sorted_keywords:
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m: f'<span style="background-color: {light_color}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        for entry in categorized:
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            matched_keywords = entry.get("matched_keywords", [])
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, matched_keywords)

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="mappa",
                categories=["MAPPA"],
                category_colors={"MAPPA": "#7c3aed"},
                highlighted_html=highlighted_html,
                parent_container=getattr(self, 'popup_mappa_import_container', None),
            )

            import_layout.addWidget(entry_frame)
            self.popup_mappa_imported_entries.append({"checkbox": cb, "text": full_text, "date": date_raw, "frame": entry_frame})

        print(f"[MOJ-ASR] Section 12 popup: Added {len(self.popup_mappa_imported_entries)} MAPPA entries")

    def _generate_attitude_behaviour(self) -> str:
        """Generate attitude/behaviour text from selections and imported data."""
        concerns = []
        negatives = []

        for key, item in self.popup_behaviour_items.items():
            if item["yes"].isChecked():
                label = item["label"].lower()
                details = item["details"].text().strip()
                if details:
                    concerns.append(f"{label} ({details})")
                else:
                    concerns.append(label)
            elif item["no"].isChecked():
                negatives.append(item["negative"])

        parts = []

        # Positive concerns
        if concerns:
            if len(concerns) == 1:
                parts.append(f"In the last 12 months there have been concerns regarding {concerns[0]}.")
            else:
                items = ", ".join(concerns[:-1]) + f" and {concerns[-1]}"
                parts.append(f"In the last 12 months there have been concerns regarding {items}.")

        # Negative statements
        if negatives:
            if len(negatives) == 1:
                parts.append(f"There has been {negatives[0]}.")
            else:
                neg_items = ", ".join(negatives[:-1]) + f" and {negatives[-1]}"
                parts.append(f"There has been {neg_items}.")

        # Additional notes
        notes = self.popup_attitude_notes.toPlainText().strip()
        if notes:
            parts.append(notes)

        # Imported notes (checked entries)
        imported_texts = []
        for entry in self.popup_behaviour_imported_entries:
            if entry["checkbox"].isChecked():
                date = entry.get("date", "")
                text = entry["text"]
                if date:
                    imported_texts.append(f"[{date}] {text}")
                else:
                    imported_texts.append(text)

        result = " ".join(parts) if parts else ""

        if imported_texts:
            if result:
                result += "\n\n--- Imported Notes ---\n" + "\n".join(imported_texts)
            else:
                result = "--- Imported Notes ---\n" + "\n".join(imported_texts)

        return result if result else "[No behaviour information provided]"

    def _create_collapsible_header(self, title: str, content_widget: QWidget, start_collapsed: bool = True):
        """Create a collapsible section with +/- toggle."""
        container = QWidget()
        container.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Maximum)
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(0, 2, 0, 2)
        container_layout.setSpacing(0)

        # Header row with +/- icon
        header = QPushButton(f"+ {title}" if start_collapsed else f"- {title}")
        header.setFixedHeight(24)
        header.setStyleSheet("""
            QPushButton {
                background: transparent;
                border: none;
                text-align: left;
                font-size: 16px;
                font-weight: 700;
                color: #1f2937;
                padding: 2px 0;
            }
            QPushButton:hover {
                color: #7c3aed;
            }
        """)
        header.setCursor(Qt.CursorShape.PointingHandCursor)

        def toggle():
            is_visible = content_widget.isVisible()
            content_widget.setVisible(not is_visible)
            header.setText(f"+ {title}" if is_visible else f"- {title}")

        header.clicked.connect(toggle)
        container_layout.addWidget(header)

        content_widget.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Maximum)
        content_widget.setVisible(not start_collapsed)
        container_layout.addWidget(content_widget)

        return container

    def _build_popup_addressing_issues(self):
        """Build addressing issues popup with structured prompts."""
        container, layout = self._create_popup_container("addressing_issues")

        text_style = """
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 8px;
                font-size: 16px;
            }
            QTextEdit:focus {
                border-color: #991b1b;
            }
        """

        slider_style = """
            QSlider::groove:horizontal {
                height: 6px;
                background: #e5e7eb;
                border-radius: 3px;
            }
            QSlider::handle:horizontal {
                width: 18px;
                margin: -6px 0;
                background: #7c3aed;
                border-radius: 9px;
            }
            QSlider::sub-page:horizontal {
                background: #7c3aed;
                border-radius: 3px;
            }
        """

        # ========== 1. Index offence work - COLLAPSIBLE ==========
        section1_content = QWidget()
        section1_layout = QVBoxLayout(section1_content)
        section1_layout.setContentsMargins(8, 4, 0, 4)
        section1_layout.setSpacing(4)

        self.popup_addr_index_options = [
            "None", "Considering", "Starting", "Engaging",
            "Well Engaged", "Almost Complete", "Complete"
        ]
        self.popup_addr_index_label = QLabel(self.popup_addr_index_options[0])
        self.popup_addr_index_label.setStyleSheet("color: #7c3aed; font-weight: 600; font-size: 16px;")
        section1_layout.addWidget(self.popup_addr_index_label)

        self.popup_addr_index_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_addr_index_slider.setRange(0, len(self.popup_addr_index_options) - 1)
        self.popup_addr_index_slider.setValue(0)
        self.popup_addr_index_slider.setStyleSheet(slider_style)
        self.popup_addr_index_slider.valueChanged.connect(
            lambda v: self.popup_addr_index_label.setText(self.popup_addr_index_options[v])
        )
        section1_layout.addWidget(self.popup_addr_index_slider)

        self.popup_addr_index_details = QLineEdit()
        self.popup_addr_index_details.setPlaceholderText("Additional details about index offence work...")
        self.popup_addr_index_details.setStyleSheet("padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 15px;")
        section1_layout.addWidget(self.popup_addr_index_details)

        layout.addWidget(self._create_collapsible_header("1. Work to address index offence(s) and risks", section1_content, start_collapsed=True))

        # ========== 2. OT Groups & Psychology - COLLAPSIBLE ==========
        section2_content = QWidget()
        section2_layout = QVBoxLayout(section2_content)
        section2_layout.setContentsMargins(8, 4, 0, 4)
        section2_layout.setSpacing(4)

        # OT checkboxes in grid
        ot_lbl = QLabel("<b>OT Groups:</b>")
        ot_lbl.setStyleSheet("font-size: 15px;")
        section2_layout.addWidget(ot_lbl)

        ot_grid_widget = QWidget()
        ot_grid = QGridLayout(ot_grid_widget)
        ot_grid.setContentsMargins(0, 2, 0, 2)
        ot_grid.setSpacing(2)

        ot_items = [
            ("breakfast_club", "Breakfast club"),
            ("smoothie", "Smoothie"),
            ("cooking", "Cooking"),
            ("current_affairs", "Current affairs"),
            ("self_care", "Self care"),
            ("ot_trips", "OT trips"),
            ("music", "Music"),
            ("art", "Art"),
            ("gym", "Gym"),
            ("woodwork", "Woodwork"),
            ("horticulture", "Horticulture"),
            ("physio", "Physio"),
        ]
        self.popup_addr_ot_checkboxes = {}
        for i, (key, label) in enumerate(ot_items):
            cb = QCheckBox(label)
            cb.setStyleSheet("font-size: 14px; color: #4b5563;")
            ot_grid.addWidget(cb, i // 4, i % 4)
            self.popup_addr_ot_checkboxes[key] = cb
        section2_layout.addWidget(ot_grid_widget)

        # OT Engagement slider
        ot_eng_lbl = QLabel("Engagement:")
        ot_eng_lbl.setStyleSheet("font-size: 16px; color: #4b5563;")
        section2_layout.addWidget(ot_eng_lbl)

        ot_eng_row = QHBoxLayout()
        self.popup_addr_ot_engagement_options = ["Limited", "Mixed", "Reasonable", "Good", "Very Good", "Excellent"]
        self.popup_addr_ot_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_addr_ot_slider.setRange(0, len(self.popup_addr_ot_engagement_options) - 1)
        self.popup_addr_ot_slider.setValue(2)
        self.popup_addr_ot_slider.setStyleSheet(slider_style)
        ot_eng_row.addWidget(self.popup_addr_ot_slider, 1)

        self.popup_addr_ot_eng_label = QLabel(self.popup_addr_ot_engagement_options[2])
        self.popup_addr_ot_eng_label.setStyleSheet("font-size: 16px; font-weight: 600; color: #7c3aed;")
        self.popup_addr_ot_slider.valueChanged.connect(
            lambda v: self.popup_addr_ot_eng_label.setText(self.popup_addr_ot_engagement_options[v])
        )
        ot_eng_row.addWidget(self.popup_addr_ot_eng_label)
        section2_layout.addLayout(ot_eng_row)
        section2_layout.addSpacing(6)

        # Psychology
        psych_lbl = QLabel("<b>Psychology:</b>")
        psych_lbl.setStyleSheet("font-size: 15px;")
        section2_layout.addWidget(psych_lbl)

        psych_grid_widget = QWidget()
        psych_grid = QGridLayout(psych_grid_widget)
        psych_grid.setContentsMargins(0, 2, 0, 2)
        psych_grid.setSpacing(2)

        psych_items = [
            ("one_to_one", "1-1"),
            ("risk", "Risk"),
            ("insight", "Insight"),
            ("psychoeducation", "Psychoeducation"),
            ("managing_emotions", "Managing emotions"),
            ("drugs_alcohol", "Drugs and alcohol"),
            ("carepathway", "Care pathway"),
            ("discharge_planning", "Discharge planning"),
        ]
        self.popup_addr_psych_checkboxes = {}
        for i, (key, label) in enumerate(psych_items):
            cb = QCheckBox(label)
            cb.setStyleSheet("font-size: 14px; color: #4b5563;")
            psych_grid.addWidget(cb, i // 4, i % 4)
            self.popup_addr_psych_checkboxes[key] = cb
        section2_layout.addWidget(psych_grid_widget)

        # Psychology Engagement slider
        psych_eng_lbl = QLabel("Engagement:")
        psych_eng_lbl.setStyleSheet("font-size: 16px; color: #4b5563;")
        section2_layout.addWidget(psych_eng_lbl)

        psych_eng_row = QHBoxLayout()
        self.popup_addr_psych_engagement_options = ["Limited", "Mixed", "Reasonable", "Good", "Very Good", "Excellent"]
        self.popup_addr_psych_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_addr_psych_slider.setRange(0, len(self.popup_addr_psych_engagement_options) - 1)
        self.popup_addr_psych_slider.setValue(2)
        self.popup_addr_psych_slider.setStyleSheet(slider_style)
        psych_eng_row.addWidget(self.popup_addr_psych_slider, 1)

        self.popup_addr_psych_eng_label = QLabel(self.popup_addr_psych_engagement_options[2])
        self.popup_addr_psych_eng_label.setStyleSheet("font-size: 16px; font-weight: 600; color: #7c3aed;")
        self.popup_addr_psych_slider.valueChanged.connect(
            lambda v: self.popup_addr_psych_eng_label.setText(self.popup_addr_psych_engagement_options[v])
        )
        psych_eng_row.addWidget(self.popup_addr_psych_eng_label)
        section2_layout.addLayout(psych_eng_row)

        layout.addWidget(self._create_collapsible_header("2. Prosocial activities (OT & Psychology)", section2_content, start_collapsed=True))

        # ========== 3. Attitudes to risk factors - COLLAPSIBLE ==========
        section3_content = QWidget()
        section3_layout = QVBoxLayout(section3_content)
        section3_layout.setContentsMargins(8, 4, 0, 4)
        section3_layout.setSpacing(4)

        # Risk factors in a scroll area with fixed max height
        risk_scroll = QScrollArea()
        risk_scroll.setWidgetResizable(True)
        risk_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        risk_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        risk_scroll.setMaximumHeight(200)
        risk_scroll.setStyleSheet("QScrollArea { background: transparent; border: 1px solid #e5e7eb; border-radius: 6px; }")

        risk_content = QWidget()
        risk_content.setStyleSheet("background: #fafafa;")
        risk_layout = QVBoxLayout(risk_content)
        risk_layout.setContentsMargins(8, 8, 8, 8)
        risk_layout.setSpacing(4)

        # Attitude slider options
        attitude_options = ["Avoids", "Limited understanding", "Some understanding", "Good understanding", "Fully understands"]

        risk_factors = [
            ("violence_others", "Violence to others"),
            ("violence_property", "Violence to property"),
            ("verbal_aggression", "Verbal aggression"),
            ("substance_misuse", "Substance misuse"),
            ("self_harm", "Self harm"),
            ("self_neglect", "Self neglect"),
            ("stalking", "Stalking"),
            ("threatening_behaviour", "Threatening behaviour"),
            ("sexually_inappropriate", "Sexually inappropriate behaviour"),
            ("vulnerability", "Vulnerability"),
            ("bullying_victimisation", "Bullying/victimisation"),
            ("absconding", "Absconding/AWOL"),
            ("reoffending", "Reoffending"),
        ]

        self.popup_addr_risk_factors = {}
        for key, label in risk_factors:
            # Container for each risk factor
            rf_container = QWidget()
            rf_layout = QVBoxLayout(rf_container)
            rf_layout.setContentsMargins(0, 0, 0, 2)
            rf_layout.setSpacing(1)

            # Checkbox
            cb = QCheckBox(label)
            cb.setStyleSheet("font-size: 14px; font-weight: 600; color: #374151; background: transparent;")
            rf_layout.addWidget(cb)

            # Slider container (hidden by default)
            slider_container = QWidget()
            slider_layout = QHBoxLayout(slider_container)
            slider_layout.setContentsMargins(20, 0, 0, 0)
            slider_layout.setSpacing(6)

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setRange(0, len(attitude_options) - 1)
            slider.setValue(2)
            slider.setStyleSheet(slider_style)
            slider_layout.addWidget(slider, 1)

            slider_label = QLabel(attitude_options[2])
            slider_label.setStyleSheet("font-size: 16px; font-weight: 600; color: #7c3aed; background: transparent;")
            slider.valueChanged.connect(lambda v, lbl=slider_label: lbl.setText(attitude_options[v]))
            slider_layout.addWidget(slider_label)
            slider_layout.addStretch()

            slider_container.hide()
            rf_layout.addWidget(slider_container)

            # Connect checkbox to show/hide slider
            cb.toggled.connect(slider_container.setVisible)

            # Connect checkbox to sync with Section 9 risk factors
            cb.toggled.connect(lambda checked, k=key: self._sync_risk_to_section9(k, checked))

            risk_layout.addWidget(rf_container)

            self.popup_addr_risk_factors[key] = {
                "checkbox": cb,
                "slider": slider,
                "slider_label": slider_label,
                "slider_container": slider_container,
                "label": label
            }

        risk_scroll.setWidget(risk_content)
        section3_layout.addWidget(risk_scroll)

        layout.addWidget(self._create_collapsible_header("3. Attitudes to risk factors", section3_content, start_collapsed=True))

        # ========== 4. Treatment for risk factors - COLLAPSIBLE ==========
        section4_content = QWidget()
        section4_layout = QVBoxLayout(section4_content)
        section4_layout.setContentsMargins(8, 4, 0, 4)
        section4_layout.setSpacing(4)

        # Container for dynamic risk factor treatment UI
        self.treatment_container = QWidget()
        treatment_main_layout = QVBoxLayout(self.treatment_container)
        treatment_main_layout.setContentsMargins(0, 0, 0, 0)
        treatment_main_layout.setSpacing(6)

        # Radio buttons for selected risk factors (populated dynamically)
        self.treatment_risk_radios_container = QWidget()
        self.treatment_risk_radios_layout = QHBoxLayout(self.treatment_risk_radios_container)
        self.treatment_risk_radios_layout.setContentsMargins(0, 0, 0, 0)
        self.treatment_risk_radios_layout.setSpacing(4)
        self.treatment_risk_radio_group = QButtonGroup(self)
        self.treatment_risk_radios = {}
        treatment_main_layout.addWidget(self.treatment_risk_radios_container)

        # Shared concern section for current risk factor (appears below radios when any treatment is ticked)
        self.risk_concern_container = QWidget()
        self.risk_concern_container.setStyleSheet("background: #fef2f2; border: 1px solid #fecaca; border-radius: 6px;")
        risk_concern_layout = QVBoxLayout(self.risk_concern_container)
        risk_concern_layout.setContentsMargins(10, 8, 10, 8)
        risk_concern_layout.setSpacing(6)

        self.risk_concern_title = QLabel("<b>Remaining concerns for this risk factor:</b>")
        self.risk_concern_title.setStyleSheet("font-size: 15px; color: #991b1b; background: transparent; border: none;")
        risk_concern_layout.addWidget(self.risk_concern_title)

        concern_row = QHBoxLayout()
        self.concern_options = ["Nil", "Minor", "Moderate", "Significant", "High"]
        self.risk_concern_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.risk_concern_slider.setRange(0, len(self.concern_options) - 1)
        self.risk_concern_slider.setValue(0)
        self.risk_concern_slider.setStyleSheet(slider_style.replace("#7c3aed", "#dc2626"))
        concern_row.addWidget(self.risk_concern_slider, 1)
        self.risk_concern_label = QLabel(self.concern_options[0])
        self.risk_concern_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #dc2626; background: transparent; border: none;")
        self.risk_concern_slider.valueChanged.connect(lambda v: self.risk_concern_label.setText(self.concern_options[v]))
        concern_row.addWidget(self.risk_concern_label)
        risk_concern_layout.addLayout(concern_row)

        # Concern details text box
        self.risk_concern_details = QLineEdit()
        self.risk_concern_details.setPlaceholderText("Specify remaining concerns...")
        self.risk_concern_details.setStyleSheet("font-size: 14px; padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; background: white;")
        self.risk_concern_details.hide()
        self.risk_concern_slider.valueChanged.connect(lambda v: self.risk_concern_details.setVisible(v > 0))
        risk_concern_layout.addWidget(self.risk_concern_details)

        self.risk_concern_container.hide()
        treatment_main_layout.addWidget(self.risk_concern_container)

        # Placeholder when no risk factors selected
        self.treatment_placeholder = QLabel("Select risk factors in Section 3 above to configure treatments")
        self.treatment_placeholder.setStyleSheet("color: #6b7280; font-style: italic; font-size: 15px; padding: 8px;")
        treatment_main_layout.addWidget(self.treatment_placeholder)

        # Treatment options container (shown when a risk radio is selected)
        self.treatment_options_container = QWidget()
        self.treatment_options_container.setStyleSheet("background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 6px;")
        treatment_options_layout = QVBoxLayout(self.treatment_options_container)
        treatment_options_layout.setContentsMargins(10, 10, 10, 10)
        treatment_options_layout.setSpacing(6)

        # Treatment types with effectiveness slider only
        treatment_types = [
            ("medication", "Medication"),
            ("psych_1to1", "Psychology 1-1"),
            ("psych_groups", "Psychology groups"),
            ("nursing", "Nursing support"),
            ("ot_support", "OT support"),
            ("social_work", "Social Work"),
        ]

        effectiveness_options = ["Nil", "Minimal", "Some", "Reasonable", "Good", "Very Good", "Excellent"]

        self.treatment_widgets = {}
        for tx_key, tx_label in treatment_types:
            tx_frame = QFrame()
            tx_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            tx_layout = QVBoxLayout(tx_frame)
            tx_layout.setContentsMargins(4, 4, 4, 4)
            tx_layout.setSpacing(4)

            # Checkbox for treatment
            tx_cb = QCheckBox(tx_label)
            tx_cb.setStyleSheet("font-size: 15px; font-weight: 600; color: #374151;")
            tx_layout.addWidget(tx_cb)

            # Effectiveness slider container (hidden until checkbox ticked)
            tx_sliders = QWidget()
            tx_sliders_layout = QVBoxLayout(tx_sliders)
            tx_sliders_layout.setContentsMargins(16, 2, 0, 2)
            tx_sliders_layout.setSpacing(2)

            # Effectiveness slider
            eff_lbl = QLabel("Effectiveness:")
            eff_lbl.setStyleSheet("font-size: 16px; color: #4b5563;")
            tx_sliders_layout.addWidget(eff_lbl)
            eff_row = QHBoxLayout()
            tx_eff_slider = NoWheelSlider(Qt.Orientation.Horizontal)
            tx_eff_slider.setRange(0, len(effectiveness_options) - 1)
            tx_eff_slider.setValue(0)
            tx_eff_slider.setStyleSheet(slider_style)
            eff_row.addWidget(tx_eff_slider, 1)
            tx_eff_label = QLabel(effectiveness_options[0])
            tx_eff_label.setStyleSheet("font-size: 16px; font-weight: 600; color: #059669;")
            tx_eff_slider.valueChanged.connect(lambda v, lbl=tx_eff_label: lbl.setText(effectiveness_options[v]))
            eff_row.addWidget(tx_eff_label)
            tx_sliders_layout.addLayout(eff_row)

            tx_sliders.hide()
            tx_cb.toggled.connect(tx_sliders.setVisible)
            tx_cb.toggled.connect(self._update_risk_concern_visibility)
            tx_layout.addWidget(tx_sliders)

            treatment_options_layout.addWidget(tx_frame)

            self.treatment_widgets[tx_key] = {
                "checkbox": tx_cb,
                "eff_slider": tx_eff_slider,
                "eff_label": tx_eff_label,
                "sliders_container": tx_sliders,
                "label": tx_label
            }

        self.treatment_options_container.hide()
        treatment_main_layout.addWidget(self.treatment_options_container)

        # Data storage: {risk_factor_key: {treatment_key: {checkbox, eff, concern, details}}}
        self.treatment_data = {}
        self.current_treatment_risk = None

        # Connect risk factor checkboxes to update treatment radios
        for key, rf_data in self.popup_addr_risk_factors.items():
            rf_data["checkbox"].toggled.connect(self._update_treatment_risk_radios)

        section4_layout.addWidget(self.treatment_container)

        layout.addWidget(self._create_collapsible_header("4. Treatment for risk factors", section4_content, start_collapsed=True))

        # ========== 5. Relapse prevention - COLLAPSIBLE ==========
        section5_content = QWidget()
        section5_layout = QVBoxLayout(section5_content)
        section5_layout.setContentsMargins(8, 4, 0, 4)
        section5_layout.setSpacing(4)

        self.relapse_options = ["Not started", "Just started", "Ongoing", "Significant\nprogression", "Almost\ncompleted", "Completed"]
        relapse_row = QHBoxLayout()
        self.popup_addr_relapse_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_addr_relapse_slider.setRange(0, len(self.relapse_options) - 1)
        self.popup_addr_relapse_slider.setValue(0)
        self.popup_addr_relapse_slider.setStyleSheet(slider_style)
        relapse_row.addWidget(self.popup_addr_relapse_slider, 1)
        self.popup_addr_relapse_label = QLabel(self.relapse_options[0])
        self.popup_addr_relapse_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #7c3aed;")
        self.popup_addr_relapse_slider.valueChanged.connect(lambda v: self.popup_addr_relapse_label.setText(self.relapse_options[v]))
        relapse_row.addWidget(self.popup_addr_relapse_label)
        section5_layout.addLayout(relapse_row)

        layout.addWidget(self._create_collapsible_header("5. Relapse prevention", section5_content, start_collapsed=True))

        layout.addStretch()
        self._add_send_button(layout, "addressing_issues", self._generate_addressing_issues)
        # Collect all OT, Psychology checkboxes, risk factor widgets, and treatment widgets for preview updates
        ot_cbs = list(self.popup_addr_ot_checkboxes.values())
        psych_cbs = list(self.popup_addr_psych_checkboxes.values())
        risk_widgets = []
        for rf_data in self.popup_addr_risk_factors.values():
            risk_widgets.append(rf_data["checkbox"])
            risk_widgets.append(rf_data["slider"])
        treatment_widgets_list = []
        for tx_data in self.treatment_widgets.values():
            treatment_widgets_list.append(tx_data["checkbox"])
            treatment_widgets_list.append(tx_data["eff_slider"])
        self._connect_preview_updates("addressing_issues", [
            self.popup_addr_index_slider, self.popup_addr_index_details,
            self.popup_addr_ot_slider, self.popup_addr_psych_slider,
            *ot_cbs, *psych_cbs, *risk_widgets, *treatment_widgets_list,
            self.risk_concern_slider, self.risk_concern_details,
            self.popup_addr_relapse_slider
        ])

    def _update_risk_concern_visibility(self):
        """Show/hide the shared risk concern section based on whether any treatment is ticked."""
        any_ticked = any(tx["checkbox"].isChecked() for tx in self.treatment_widgets.values())
        self.risk_concern_container.setVisible(any_ticked)

    def _update_treatment_risk_radios(self):
        """Update treatment risk radio buttons based on selected risk factors in Section 3."""
        # Save current state before clearing
        if self.current_treatment_risk:
            self._save_treatment_state(self.current_treatment_risk)

        # Clear existing radios
        for radio in self.treatment_risk_radios.values():
            self.treatment_risk_radio_group.removeButton(radio)
            radio.deleteLater()
        self.treatment_risk_radios.clear()

        # Clear layout
        while self.treatment_risk_radios_layout.count():
            item = self.treatment_risk_radios_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Get selected risk factors
        selected_risks = []
        risk_labels = {
            "violence_others": "Violence to others",
            "violence_property": "Violence to property",
            "verbal_aggression": "Verbal aggression",
            "substance_misuse": "Substance misuse",
            "self_harm": "Self harm",
            "self_neglect": "Self neglect",
            "stalking": "Stalking",
            "threatening_behaviour": "Threatening behaviour",
            "sexually_inappropriate": "Sexually inappropriate",
            "vulnerability": "Vulnerability",
            "bullying_victimisation": "Bullying/victimisation",
            "absconding": "Absconding/AWOL",
            "reoffending": "Reoffending",
        }

        for key, rf_data in self.popup_addr_risk_factors.items():
            if rf_data["checkbox"].isChecked():
                selected_risks.append((key, risk_labels.get(key, key)))

        if not selected_risks:
            self.treatment_placeholder.show()
            self.treatment_options_container.hide()
            self.current_treatment_risk = None
            return

        self.treatment_placeholder.hide()

        # Create radio buttons for each selected risk
        for key, label in selected_risks:
            radio = QRadioButton(label)
            radio.setStyleSheet("""
                QRadioButton { font-size: 14px; font-weight: 600; color: #374151; padding: 4px 8px; }
                QRadioButton:checked { color: #7c3aed; }
            """)
            radio.toggled.connect(lambda checked, k=key: self._on_treatment_risk_changed(k) if checked else None)
            self.treatment_risk_radio_group.addButton(radio)
            self.treatment_risk_radios[key] = radio
            self.treatment_risk_radios_layout.addWidget(radio)

        self.treatment_risk_radios_layout.addStretch()

        # Select first radio by default
        if selected_risks:
            first_key = selected_risks[0][0]
            self.treatment_risk_radios[first_key].setChecked(True)

    def _on_treatment_risk_changed(self, risk_key):
        """Handle switching between risk factor treatment tabs."""
        # Save current state
        if self.current_treatment_risk and self.current_treatment_risk != risk_key:
            self._save_treatment_state(self.current_treatment_risk)

        self.current_treatment_risk = risk_key
        self.treatment_options_container.show()

        # Restore state for new risk
        self._restore_treatment_state(risk_key)

    def _save_treatment_state(self, risk_key):
        """Save current treatment widget state for a risk factor."""
        if risk_key not in self.treatment_data:
            self.treatment_data[risk_key] = {"treatments": {}, "concern": 0, "concern_details": ""}

        # Save treatment checkboxes and effectiveness
        for tx_key, tx_widgets in self.treatment_widgets.items():
            self.treatment_data[risk_key]["treatments"][tx_key] = {
                "checked": tx_widgets["checkbox"].isChecked(),
                "effectiveness": tx_widgets["eff_slider"].value(),
            }

        # Save shared concern for this risk factor
        self.treatment_data[risk_key]["concern"] = self.risk_concern_slider.value()
        self.treatment_data[risk_key]["concern_details"] = self.risk_concern_details.text()

    def _restore_treatment_state(self, risk_key):
        """Restore treatment widget state for a risk factor."""
        saved_data = self.treatment_data.get(risk_key, {"treatments": {}, "concern": 0, "concern_details": ""})
        treatments_data = saved_data.get("treatments", {})

        any_checked = False
        for tx_key, tx_widgets in self.treatment_widgets.items():
            tx_saved = treatments_data.get(tx_key, {})

            # Block signals to prevent triggering updates
            tx_widgets["checkbox"].blockSignals(True)
            tx_widgets["eff_slider"].blockSignals(True)

            checked = tx_saved.get("checked", False)
            tx_widgets["checkbox"].setChecked(checked)
            tx_widgets["eff_slider"].setValue(tx_saved.get("effectiveness", 0))

            # Update visibility
            tx_widgets["sliders_container"].setVisible(checked)

            # Update labels
            effectiveness_options = ["Nil", "Minimal", "Some", "Reasonable", "Good", "Very Good", "Excellent"]
            tx_widgets["eff_label"].setText(effectiveness_options[tx_saved.get("effectiveness", 0)])

            # Unblock signals
            tx_widgets["checkbox"].blockSignals(False)
            tx_widgets["eff_slider"].blockSignals(False)

            if checked:
                any_checked = True

        # Restore shared concern slider
        self.risk_concern_slider.blockSignals(True)
        concern_val = saved_data.get("concern", 0)
        self.risk_concern_slider.setValue(concern_val)
        self.risk_concern_label.setText(self.concern_options[concern_val])
        self.risk_concern_details.setText(saved_data.get("concern_details", ""))
        self.risk_concern_details.setVisible(concern_val > 0)
        self.risk_concern_slider.blockSignals(False)

        # Show/hide concern container based on whether any treatment is checked
        self.risk_concern_container.setVisible(any_checked)

    def _generate_addressing_issues(self) -> str:
        """Generate addressing issues narrative text from slider and fields."""
        p = self._get_pronouns()
        parts = []

        # Helper function to format lists with "and" before last item
        def format_list(items):
            if not items:
                return ""
            if len(items) == 1:
                return items[0]
            if len(items) == 2:
                return f"{items[0]} and {items[1]}"
            return ", ".join(items[:-1]) + f" and {items[-1]}"

        # 1. Index offence work - from slider (grammar-aware)
        idx_val = self.popup_addr_index_slider.value()
        # Use pronoun-aware verb conjugation
        idx_options_narrative = {
            0: None,  # None - skip
            1: f"{p['is']} considering engaging with work",
            2: f"{p['has']} started work",
            3: f"{p['is']} engaging with work",
            4: f"{p['is']} well engaged with work",
            5: f"{p['has']} almost completed work",
            6: f"{p['has']} completed work"
        }
        if idx_val > 0:
            idx_text = f"{p['subj']} {idx_options_narrative[idx_val]} to address the index offence and associated risks"
            details = self.popup_addr_index_details.text().strip()
            if details:
                idx_text += f", including {details}"
            parts.append(idx_text + ".")

        # 2. OT Groups - narrative output
        ot_display_names = {
            "breakfast_club": "breakfast club",
            "smoothie": "smoothie groups",
            "cooking": "cooking",
            "current_affairs": "current affairs",
            "self_care": "self care",
            "ot_trips": "OT trips",
            "music": "music",
            "art": "art",
            "gym": "gym",
            "woodwork": "woodwork",
            "horticulture": "horticulture",
            "physio": "physio",
        }
        ot_selected = [ot_display_names[key] for key, cb in self.popup_addr_ot_checkboxes.items() if cb.isChecked()]

        if ot_selected:
            ot_list = format_list(ot_selected)
            ot_text = f"{p['subj']} {p['engages']} in prosocial groups, such as in OT - {ot_list}."

            # Engagement level narrative
            eng_val = self.popup_addr_ot_slider.value()
            eng_narrative = {
                0: "limited however",
                1: "mixed",
                2: "reasonably good",
                3: "good",
                4: "very good",
                5: "excellent"
            }
            ot_text += f" Overall engagement in these groups is {eng_narrative[eng_val]}."
            parts.append(ot_text)

        # 2b. Psychology Engagement - narrative output
        psych_display_names = {
            "one_to_one": "1-1s",
            "risk": "risk",
            "insight": "insight",
            "psychoeducation": "psychoeducation",
            "managing_emotions": "managing emotions",
            "drugs_alcohol": "drugs and alcohol",
            "carepathway": "care pathway",
            "discharge_planning": "discharge planning",
        }
        psych_selected = [(key, psych_display_names[key]) for key, cb in self.popup_addr_psych_checkboxes.items() if cb.isChecked()]

        if psych_selected:
            has_one_to_one = any(key == "one_to_one" for key, _ in psych_selected)
            groups = [name for key, name in psych_selected if key != "one_to_one"]

            # Build psychology sentence
            if ot_selected:
                psych_text = f"Likewise {p['subj_l']} {p['engages']} in psychology"
            else:
                psych_text = f"{p['subj']} {p['engages']} in psychology"

            psych_parts = []
            if has_one_to_one:
                psych_parts.append("utilising 1-1s")
            if groups:
                groups_list = format_list(groups)
                if has_one_to_one:
                    psych_parts.append(f"and groups/sessions into {groups_list}")
                else:
                    psych_parts.append(f"groups/sessions into {groups_list}")

            if psych_parts:
                psych_text += " " + " ".join(psych_parts) + "."
            else:
                psych_text += "."

            # Add psychology engagement slider
            psych_eng_val = self.popup_addr_psych_slider.value()
            psych_eng_narrative = {
                0: "limited however",
                1: "mixed",
                2: "reasonably good",
                3: "good",
                4: "very good",
                5: "excellent"
            }
            psych_text += f" Overall engagement in psychology is {psych_eng_narrative[psych_eng_val]}."
            parts.append(psych_text)

        # 3. Attitudes to risk factors - grouped by understanding level, most positive first
        # Collect selected risk factors with their slider values
        risk_by_level = {
            4: [],  # Fully understands
            3: [],  # Good understanding
            2: [],  # Some understanding
            1: [],  # Limited understanding
            0: [],  # Avoids
        }
        risk_display_names = {
            "violence_others": f"violence to others",
            "violence_property": f"violence to property",
            "verbal_aggression": f"verbal aggression",
            "substance_misuse": f"substance misuse",
            "self_harm": f"self harm",
            "self_neglect": f"self neglect",
            "stalking": f"stalking",
            "threatening_behaviour": f"threatening behaviour",
            "sexually_inappropriate": f"sexually inappropriate behaviour",
            "vulnerability": f"vulnerability",
            "bullying_victimisation": f"bullying/victimisation",
            "absconding": f"absconding/AWOL",
            "reoffending": f"reoffending",
        }

        for key, rf_data in self.popup_addr_risk_factors.items():
            if rf_data["checkbox"].isChecked():
                level = rf_data["slider"].value()
                risk_by_level[level].append(risk_display_names[key])

        # Build risk factors narrative - most positive first
        risk_parts = []
        has_any_risk = any(risks for risks in risk_by_level.values())

        if has_any_risk:
            # Level descriptions for narrative (avoiding repetition of "understanding")
            level_intros = {
                4: "full understanding of",
                3: "good understanding of",
                2: "some understanding of",
                1: "limited understanding of",
                0: "avoids discussing",
            }
            level_connectors = {
                4: "",
                3: "",
                2: "but only",
                1: "but only",
                0: "and",
            }

            # Build sentence parts from most positive (4) to least positive (0)
            sentence_parts = []
            first_part = True
            for level in [4, 3, 2, 1, 0]:
                risks = risk_by_level[level]
                if risks:
                    risk_list = format_list(risks)
                    if first_part:
                        sentence_parts.append(f"{level_intros[level]} {p['pos_l']} risk of {risk_list}")
                        first_part = False
                    else:
                        connector = level_connectors[level]
                        if level == 0:
                            sentence_parts.append(f"{connector} {level_intros[level]} {risk_list}")
                        else:
                            sentence_parts.append(f"{connector} {level_intros[level]} {risk_list}")

            if sentence_parts:
                risk_text = f"{p['pos']} attitudes to risk factors reveal " + ", ".join(sentence_parts) + "."
                parts.append(risk_text)

        # 4. Treatment for risk factors - save current state first
        if self.current_treatment_risk:
            self._save_treatment_state(self.current_treatment_risk)

        # Collect all treatment data
        treatment_labels = {
            "medication": "medication",
            "psych_1to1": "1-1 psychology",
            "psych_groups": "psychology groups",
            "nursing": "nursing support",
            "ot_support": "OT support",
            "social_work": "social work",
        }
        risk_labels = {
            "violence_others": "violence to others",
            "violence_property": "violence to property",
            "verbal_aggression": "verbal aggression",
            "substance_misuse": "substance misuse",
            "self_harm": "self harm",
            "self_neglect": "self neglect",
            "stalking": "stalking",
            "threatening_behaviour": "threatening behaviour",
            "sexually_inappropriate": "sexually inappropriate behaviour",
            "vulnerability": "vulnerability",
            "bullying_victimisation": "bullying/victimisation",
            "absconding": "absconding/AWOL",
            "reoffending": "reoffending",
        }
        effectiveness_levels = ["nil", "minimal", "some", "reasonable", "good", "very good", "excellent"]
        concern_levels = ["nil", "minor", "moderate", "significant", "high"]

        # Build treatment entries from new data structure
        # Structure: {risk_key: {treatments: {tx_key: {checked, effectiveness}}, concern, concern_details}}
        treatment_entries = []
        risk_concerns = {}  # Store concern per risk factor
        for risk_key, risk_data in self.treatment_data.items():
            treatments = risk_data.get("treatments", {})
            concern = risk_data.get("concern", 0)
            concern_details = risk_data.get("concern_details", "")
            risk_concerns[risk_key] = {"level": concern, "details": concern_details}

            for tx_key, tx_data in treatments.items():
                if tx_data.get("checked", False):
                    treatment_entries.append({
                        "risk_key": risk_key,
                        "risk_label": risk_labels.get(risk_key, risk_key),
                        "tx_key": tx_key,
                        "tx_label": treatment_labels.get(tx_key, tx_key),
                        "effectiveness": tx_data.get("effectiveness", 0),
                    })

        if treatment_entries:
            # Sort by effectiveness (highest first)
            treatment_entries.sort(key=lambda x: x["effectiveness"], reverse=True)

            # Group by risk factor for treatment sentence
            risks_with_treatments = {}
            for entry in treatment_entries:
                rk = entry["risk_key"]
                if rk not in risks_with_treatments:
                    risks_with_treatments[rk] = []
                risks_with_treatments[rk].append(entry)

            if len(risks_with_treatments) == 1:
                # Single risk factor - simpler output
                risk_key = list(risks_with_treatments.keys())[0]
                entries = risks_with_treatments[risk_key]
                tx_list = format_list([e["tx_label"] for e in entries])
                risk_label = risk_labels.get(risk_key, risk_key)

                # Get highest effectiveness
                max_eff = max(e["effectiveness"] for e in entries)
                eff_text = effectiveness_levels[max_eff]

                tx_text = f"{p['subj']} {p['has']} engaged in {tx_list} for {p['pos_l']} risk of {risk_label}"
                if max_eff == 0:
                    tx_text += " with no effectiveness to date"
                else:
                    tx_text += f" with {eff_text} effectiveness"

                # Get concern for this risk factor
                risk_concern = risk_concerns.get(risk_key, {"level": 0, "details": ""})
                concern_level = risk_concern["level"]
                concern_details = risk_concern["details"]

                if concern_level == 0:
                    tx_text += " and no remaining concerns"
                else:
                    concern_text = concern_levels[concern_level]
                    if concern_details:
                        tx_text += f" and {concern_text} remaining concerns ({concern_details})"
                    else:
                        tx_text += f" and {concern_text} remaining concerns"

                parts.append(tx_text + ".")

            else:
                # Multiple risk factors - complex output
                # First sentence: what treatment for what
                risk_treatment_parts = []
                for risk_key, entries in risks_with_treatments.items():
                    risk_label = risk_labels.get(risk_key, risk_key)
                    tx_list = format_list([e["tx_label"] for e in entries])
                    risk_treatment_parts.append(f"{risk_label} ({tx_list})")

                tx_text = f"{p['subj']} {p['has']} engaged in treatment for " + format_list(risk_treatment_parts) + "."
                parts.append(tx_text)

                # Second sentence: effectiveness - get max effectiveness per risk factor
                risk_max_eff = {}
                for risk_key, entries in risks_with_treatments.items():
                    max_eff = max(e["effectiveness"] for e in entries)
                    risk_max_eff[risk_key] = max_eff

                # Group risks by their max effectiveness level
                eff_by_level = {}
                for risk_key, eff in risk_max_eff.items():
                    if eff not in eff_by_level:
                        eff_by_level[eff] = []
                    eff_by_level[eff].append(risk_labels.get(risk_key, risk_key))

                # Build effectiveness sentence - highest level first, then "less so" for lower levels
                sorted_levels = sorted(eff_by_level.keys(), reverse=True)
                if sorted_levels:
                    eff_parts = []
                    first_part = True
                    for eff_level in sorted_levels:
                        risks = eff_by_level[eff_level]
                        eff_text = effectiveness_levels[eff_level]
                        if first_part:
                            # First (highest) level - state the effectiveness
                            if eff_level == 0:
                                eff_parts.append(f"of no effectiveness for {format_list(risks)}")
                            else:
                                eff_parts.append(f"of {eff_text} effectiveness for {format_list(risks)}")
                            first_part = False
                        else:
                            # Lower levels - use "less so" or "of no effectiveness"
                            if eff_level == 0:
                                eff_parts.append(f"of no effectiveness for {format_list(risks)}")
                            else:
                                eff_parts.append(f"less so for {format_list(risks)}")

                    if eff_parts:
                        parts.append(f"The treatment has been " + " but ".join(eff_parts) + ".")

                # Third sentence: concerns - one per risk factor
                risks_with_concerns = {rk: risk_concerns[rk] for rk in risks_with_treatments.keys() if risk_concerns.get(rk, {}).get("level", 0) > 0}
                risks_no_concerns = [rk for rk in risks_with_treatments.keys() if risk_concerns.get(rk, {}).get("level", 0) == 0]

                if risks_with_concerns:
                    # Group by concern level
                    concern_by_level = {}
                    for rk, data in risks_with_concerns.items():
                        level = data["level"]
                        if level not in concern_by_level:
                            concern_by_level[level] = {"risks": [], "details": []}
                        concern_by_level[level]["risks"].append(risk_labels.get(rk, rk))
                        if data["details"]:
                            concern_by_level[level]["details"].append(data["details"])

                    # Build concern sentence - highest concern level first
                    concern_parts = []
                    for c_level in sorted(concern_by_level.keys(), reverse=True):
                        c_data = concern_by_level[c_level]
                        level_text = concern_levels[c_level].capitalize()
                        risk_list = format_list(c_data["risks"])
                        if c_data["details"]:
                            concern_parts.append(f"{level_text} concerns remain around {risk_list}, specifically {'; '.join(c_data['details'])}")
                        else:
                            concern_parts.append(f"{level_text} concerns remain around {risk_list}")

                    if concern_parts:
                        parts.append(". ".join(concern_parts) + ".")

                if risks_no_concerns:
                    no_concern_list = format_list([risk_labels.get(rk, rk) for rk in risks_no_concerns])
                    parts.append(f"There are no remaining concerns around {no_concern_list}.")

        # 5. Relapse prevention - narrative output from slider
        relapse_val = self.popup_addr_relapse_slider.value()
        relapse_narratives = {
            0: f"Relapse prevention work has not yet started but is planned as {p['subj_l']} {p['engages']} further in the care pathway.",
            1: f"{p['subj']} {p['has']} just started relapse prevention work.",
            2: f"{p['subj']} {p['is']} undertaking ongoing relapse prevention work.",
            3: f"{p['subj']} {p['has']} made significant progression in relapse prevention work.",
            4: f"{p['subj']} {p['has']} almost completed relapse prevention work.",
            5: f"{p['subj']} {p['has']} completed relapse prevention work."
        }
        parts.append(relapse_narratives[relapse_val])

        return " ".join(parts) if parts else "[No information provided]"

    def _build_popup_patient_attitude(self):
        """Build patient attitude popup - Understanding & Compliance grid plus Offending Behaviour."""
        container, layout = self._create_popup_container_with_imports("patient_attitude")

        combo_style = """
            QComboBox {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 15px;
                min-width: 90px;
            }
            QComboBox:focus { border-color: #991b1b; }
            QComboBox::drop-down { border: none; width: 20px; }
            QComboBox::down-arrow { image: none; border-left: 4px solid transparent; border-right: 4px solid transparent; border-top: 5px solid #6b7280; }
        """

        text_style = """
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 8px;
                font-size: 16px;
            }
            QTextEdit:focus { border-color: #991b1b; }
        """

        # ================================================================
        # SECTION 1: UNDERSTANDING & COMPLIANCE GRID
        # ================================================================
        section1_lbl = QLabel("<b>1. Understanding & Compliance with Treatment:</b>")
        section1_lbl.setWordWrap(True)
        layout.addWidget(section1_lbl)

        grid = QGridLayout()
        grid.setSpacing(6)

        # Headers
        header_treatment = QLabel("Treatment")
        header_treatment.setStyleSheet("font-size: 15px; font-weight: 600; color: #374151;")
        header_understanding = QLabel("Understanding")
        header_understanding.setStyleSheet("font-size: 15px; font-weight: 600; color: #374151;")
        header_compliance = QLabel("Compliance")
        header_compliance.setStyleSheet("font-size: 15px; font-weight: 600; color: #374151;")

        grid.addWidget(header_treatment, 0, 0)
        grid.addWidget(header_understanding, 0, 1)
        grid.addWidget(header_compliance, 0, 2)

        # Treatment rows
        self.popup_att_treatments = {}
        treatment_names = ["Medical", "Nursing", "Psychology", "OT", "Social Work"]
        understanding_options = ["Select...", "good", "fair", "poor"]
        compliance_options = ["Select...", "full", "reasonable", "partial", "nil"]

        for i, name in enumerate(treatment_names, 1):
            key = name.lower().replace(" ", "_")

            lbl = QLabel(name)
            lbl.setStyleSheet("font-size: 15px; color: #374151;")
            grid.addWidget(lbl, i, 0)

            understanding = QComboBox()
            understanding.addItems(understanding_options)
            understanding.setStyleSheet(combo_style)
            understanding.currentIndexChanged.connect(lambda _, k="patient_attitude": self._update_preview(k))
            grid.addWidget(understanding, i, 1)

            compliance = QComboBox()
            compliance.addItems(compliance_options)
            compliance.setStyleSheet(combo_style)
            compliance.currentIndexChanged.connect(lambda _, k="patient_attitude": self._update_preview(k))
            grid.addWidget(compliance, i, 2)

            self.popup_att_treatments[key] = {
                "understanding": understanding,
                "compliance": compliance
            }

        layout.addLayout(grid)
        layout.addSpacing(16)

        # ================================================================
        # SECTION 2: OFFENDING BEHAVIOUR
        # ================================================================
        section2_lbl = QLabel("<b>2. Offending Behaviour - Insight & Responsibility:</b>")
        section2_lbl.setWordWrap(True)
        layout.addWidget(section2_lbl)

        # Insight into offending slider
        layout.addWidget(QLabel("Insight into offending:"))
        insight_row = QHBoxLayout()
        self.popup_att_offending_insight_options = ["Nil", "Limited", "Partial", "Good", "Full"]
        self.popup_att_offending_insight_label = QLabel(self.popup_att_offending_insight_options[2])
        self.popup_att_offending_insight_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        self.popup_att_offending_insight_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_att_offending_insight_slider.setRange(0, len(self.popup_att_offending_insight_options) - 1)
        self.popup_att_offending_insight_slider.setValue(2)
        self.popup_att_offending_insight_slider.setStyleSheet("""
            QSlider::groove:horizontal { height: 6px; background: #e5e7eb; border-radius: 3px; }
            QSlider::handle:horizontal { width: 16px; margin: -5px 0; background: #991b1b; border-radius: 8px; }
            QSlider::sub-page:horizontal { background: #991b1b; border-radius: 3px; }
        """)
        self.popup_att_offending_insight_slider.valueChanged.connect(
            lambda v: self.popup_att_offending_insight_label.setText(self.popup_att_offending_insight_options[v])
        )
        self.popup_att_offending_insight_slider.valueChanged.connect(lambda: self._update_preview("patient_attitude"))
        insight_row.addWidget(self.popup_att_offending_insight_slider, 1)
        insight_row.addWidget(self.popup_att_offending_insight_label)
        layout.addLayout(insight_row)

        # Accepts responsibility slider
        layout.addWidget(QLabel("Accepts responsibility:"))
        resp_row = QHBoxLayout()
        self.popup_att_responsibility_options = ["Denies", "Minimises", "Partial", "Mostly", "Full"]
        self.popup_att_responsibility_label = QLabel(self.popup_att_responsibility_options[2])
        self.popup_att_responsibility_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        self.popup_att_responsibility_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_att_responsibility_slider.setRange(0, len(self.popup_att_responsibility_options) - 1)
        self.popup_att_responsibility_slider.setValue(2)
        self.popup_att_responsibility_slider.setStyleSheet("""
            QSlider::groove:horizontal { height: 6px; background: #e5e7eb; border-radius: 3px; }
            QSlider::handle:horizontal { width: 16px; margin: -5px 0; background: #991b1b; border-radius: 8px; }
            QSlider::sub-page:horizontal { background: #991b1b; border-radius: 3px; }
        """)
        self.popup_att_responsibility_slider.valueChanged.connect(
            lambda v: self.popup_att_responsibility_label.setText(self.popup_att_responsibility_options[v])
        )
        self.popup_att_responsibility_slider.valueChanged.connect(lambda: self._update_preview("patient_attitude"))
        resp_row.addWidget(self.popup_att_responsibility_slider, 1)
        resp_row.addWidget(self.popup_att_responsibility_label)
        layout.addLayout(resp_row)

        # Victim empathy slider
        layout.addWidget(QLabel("Victim empathy:"))
        empathy_row = QHBoxLayout()
        self.popup_att_empathy_options = ["Nil", "Limited", "Developing", "Good", "Full"]
        self.popup_att_empathy_label = QLabel(self.popup_att_empathy_options[2])
        self.popup_att_empathy_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        self.popup_att_empathy_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_att_empathy_slider.setRange(0, len(self.popup_att_empathy_options) - 1)
        self.popup_att_empathy_slider.setValue(2)
        self.popup_att_empathy_slider.setStyleSheet("""
            QSlider::groove:horizontal { height: 6px; background: #e5e7eb; border-radius: 3px; }
            QSlider::handle:horizontal { width: 16px; margin: -5px 0; background: #991b1b; border-radius: 8px; }
            QSlider::sub-page:horizontal { background: #991b1b; border-radius: 3px; }
        """)
        self.popup_att_empathy_slider.valueChanged.connect(
            lambda v: self.popup_att_empathy_label.setText(self.popup_att_empathy_options[v])
        )
        self.popup_att_empathy_slider.valueChanged.connect(lambda: self._update_preview("patient_attitude"))
        empathy_row.addWidget(self.popup_att_empathy_slider, 1)
        empathy_row.addWidget(self.popup_att_empathy_label)
        layout.addLayout(empathy_row)

        # Additional details text area
        layout.addSpacing(8)
        layout.addWidget(QLabel("Additional details:"))
        self.popup_att_offending_details = QTextEdit()
        self.popup_att_offending_details.setMaximumHeight(60)
        self.popup_att_offending_details.setPlaceholderText("Any additional context about insight into offending or responsibility...")
        self.popup_att_offending_details.setStyleSheet(text_style)
        self.popup_att_offending_details.textChanged.connect(lambda: self._update_preview("patient_attitude"))
        layout.addWidget(self.popup_att_offending_details)

        layout.addStretch()
        self._add_send_button(layout, "patient_attitude", self._generate_patient_attitude)

    def _generate_patient_attitude(self) -> str:
        """Generate patient attitude text from understanding/compliance grid and offending behaviour."""
        p = self._get_pronouns()
        parts = []

        # === SECTION 1: Understanding & Compliance ===
        # Medical
        med = self.popup_att_treatments["medical"]
        med_u = med["understanding"].currentText()
        med_c = med["compliance"].currentText()
        if med_u != "Select..." and med_c != "Select...":
            u_phrase = self._att_understanding_phrase(med_u, "medical", p)
            c_phrase = self._att_compliance_phrase(med_c, p)
            if u_phrase and c_phrase:
                parts.append(f"{u_phrase} {c_phrase}.")

        # Nursing
        nursing = self.popup_att_treatments["nursing"]
        nursing_u = nursing["understanding"].currentText()
        nursing_c = nursing["compliance"].currentText()
        if nursing_u != "Select..." and nursing_c != "Select...":
            phrase = self._att_nursing_phrase(nursing_u, nursing_c, p)
            if phrase:
                parts.append(phrase)

        # Psychology
        psych = self.popup_att_treatments["psychology"]
        psych_u = psych["understanding"].currentText()
        psych_c = psych["compliance"].currentText()
        if psych_u != "Select..." and psych_c != "Select...":
            phrase = self._att_psychology_phrase(psych_u, psych_c, p)
            if phrase:
                parts.append(phrase)

        # OT
        ot = self.popup_att_treatments["ot"]
        ot_u = ot["understanding"].currentText()
        ot_c = ot["compliance"].currentText()
        if ot_u != "Select..." and ot_c != "Select...":
            phrase = self._att_ot_phrase(ot_u, ot_c, p)
            if phrase:
                parts.append(phrase)

        # Social Work
        sw = self.popup_att_treatments["social_work"]
        sw_u = sw["understanding"].currentText()
        sw_c = sw["compliance"].currentText()
        if sw_u != "Select..." and sw_c != "Select...":
            phrase = self._att_social_work_phrase(sw_u, sw_c, p)
            if phrase:
                parts.append(phrase)

        # === SECTION 2: Offending Behaviour ===
        insight_val = self.popup_att_offending_insight_slider.value()
        resp_val = self.popup_att_responsibility_slider.value()
        empathy_val = self.popup_att_empathy_slider.value()

        insight_text = self.popup_att_offending_insight_options[insight_val].lower()
        resp_text = self.popup_att_responsibility_options[resp_val].lower()
        empathy_text = self.popup_att_empathy_options[empathy_val].lower()

        # Build offending behaviour sentence
        offending_parts = []

        # Insight
        if insight_val == 0:
            offending_parts.append(f"{p['subj']} {p['has']} no insight into {p['pos_l']} offending behaviour")
        elif insight_val == 1:
            offending_parts.append(f"{p['subj']} {p['has']} limited insight into {p['pos_l']} offending behaviour")
        elif insight_val == 2:
            offending_parts.append(f"{p['subj']} {p['has']} partial insight into {p['pos_l']} offending behaviour")
        elif insight_val == 3:
            offending_parts.append(f"{p['subj']} {p['has']} good insight into {p['pos_l']} offending behaviour")
        else:
            offending_parts.append(f"{p['subj']} {p['has']} full insight into {p['pos_l']} offending behaviour")

        # Responsibility
        if resp_val == 0:
            offending_parts.append(f"and denies responsibility")
        elif resp_val == 1:
            offending_parts.append(f"and minimises {p['pos_l']} responsibility")
        elif resp_val == 2:
            offending_parts.append(f"and partially accepts responsibility")
        elif resp_val == 3:
            offending_parts.append(f"and mostly accepts responsibility")
        else:
            offending_parts.append(f"and fully accepts responsibility")

        # Empathy
        if empathy_val == 0:
            offending_parts.append(f"with no victim empathy")
        elif empathy_val == 1:
            offending_parts.append(f"with limited victim empathy")
        elif empathy_val == 2:
            offending_parts.append(f"with developing victim empathy")
        elif empathy_val == 3:
            offending_parts.append(f"with good victim empathy")
        else:
            offending_parts.append(f"with full victim empathy")

        parts.append(" ".join(offending_parts) + ".")

        # Additional details
        details = self.popup_att_offending_details.toPlainText().strip()
        if details:
            parts.append(details)

        # Include checked imported entries (full text)
        if hasattr(self, 'popup_patient_attitude_imported_entries'):
            for entry in self.popup_patient_attitude_imported_entries:
                if entry.get("checkbox") and entry["checkbox"].isChecked():
                    full_text = entry.get("text", "").strip()
                    if full_text:
                        parts.append(f"\n\n[Imported Note]\n{full_text}")

        return " ".join(parts) if parts else "[Select understanding and compliance levels...]"

    def _att_understanding_phrase(self, level: str, treatment: str, p: dict) -> str:
        """Generate understanding phrase for attitude popup."""
        if level == "good":
            return f"{p['subj']} {p['has']} good understanding of {p['pos_l']} {treatment} treatment"
        elif level == "fair":
            return f"{p['subj']} {p['has']} some understanding of {p['pos_l']} {treatment} treatment"
        elif level == "poor":
            return f"{p['subj']} {p['has']} limited understanding of {p['pos_l']} {treatment} treatment"
        return ""

    def _att_compliance_phrase(self, level: str, p: dict) -> str:
        """Generate compliance phrase for attitude popup."""
        if level == "full":
            return "and compliance is full"
        elif level == "reasonable":
            return "and compliance is reasonable"
        elif level == "partial":
            return "but compliance is partial"
        elif level == "nil":
            return "and compliance is nil"
        return ""

    def _att_nursing_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural nursing phrase for attitude popup."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['engages']} well with nursing staff and {p['sees']} the need for nursing input."
        elif understanding == "good" and compliance == "partial":
            return f"{p['subj']} understands the role of nursing but engagement is inconsistent."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"{p['subj']} has some understanding of nursing care and {p['engages']} reasonably well."
        elif understanding == "fair" and compliance == "partial":
            return f"{p['subj']} has some understanding of nursing input but {p['engages']} only partially."
        elif understanding == "poor" or compliance == "nil":
            return f"{p['subj']} has limited insight into the need for nursing care and {p['does']} not engage meaningfully."
        return ""

    def _att_psychology_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural psychology phrase for attitude popup."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['engages']} in psychology sessions and sees the benefit of this work."
        elif understanding == "good" and compliance == "partial":
            return f"{p['subj']} understands the purpose of psychology but compliance with sessions is limited."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"{p['subj']} has some understanding of psychology and attends sessions regularly."
        elif understanding == "fair" and compliance == "partial":
            return f"{p['subj']} also {p['engages']} in psychology sessions but the compliance with these is limited."
        elif understanding == "poor" or compliance == "nil":
            return f"{p['subj']} has limited insight into the need for psychology and {p['does']} not engage with sessions."
        return ""

    def _att_ot_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural OT phrase for attitude popup."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"With respect to OT, {p['subj_l']} {p['engages']} well and sees the benefit of activities."
        elif understanding == "good" and compliance == "partial":
            return f"With respect to OT, {p['subj_l']} understands the purpose but engagement is inconsistent."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"With respect to OT, {p['subj_l']} has some understanding and participates in activities."
        elif understanding == "fair" and compliance == "partial":
            return f"With respect to OT, {p['subj_l']} has some insight but engagement is limited."
        elif understanding == "poor" or compliance == "nil":
            return f"With respect to OT, {p['subj_l']} {p['is']} not engaging and doesn't see the need to."
        return ""

    def _att_social_work_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural social work phrase for attitude popup."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['engages']} well with the social worker and understands {p['pos_l']} social circumstances."
        elif understanding == "good" and compliance == "partial":
            return f"{p['subj']} understands the social worker's role but engagement is inconsistent."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"{p['subj']} has some understanding of social work input and {p['engages']} when available."
        elif understanding == "fair" and compliance == "partial":
            return f"{p['subj']} occasionally {p['sees']} the social worker and {p['engages']} partially when {p['subj_l']} {p['does']} so."
        elif understanding == "poor" or compliance == "nil":
            return f"{p['subj']} has limited engagement with social work and doesn't see the relevance."
        return ""

    def _build_popup_capacity(self):
        """Build capacity popup with per-area capacity assessment and conditional actions."""
        container, layout = self._create_popup_container_with_imports("capacity")

        combo_style = """
            QComboBox {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 15px;
                min-width: 120px;
            }
            QComboBox:focus { border-color: #991b1b; }
        """

        cb_style = """
            QCheckBox {
                font-size: 15px;
                spacing: 4px;
            }
        """

        layout.addWidget(QLabel("<b>Capacity Assessment by Area:</b>"))
        layout.addSpacing(4)

        # Store capacity areas and their widgets
        self.popup_cap_areas = {}
        capacity_options = ["Select...", "Has capacity", "Lacks capacity"]

        # Define areas and their action options
        areas_config = {
            "residence": {
                "label": "Residence",
                "actions": ["Best Interest", "IMCA", "DoLS", "COP"]
            },
            "medication": {
                "label": "Medication",
                "actions": "soad"  # Special case - SOAD Yes/No radios
            },
            "finances": {
                "label": "Finances",
                "actions": "finance_special"  # Special case - checkboxes + radio group
            },
            "leave": {
                "label": "Leave",
                "actions": ["Best Interest", "IMCA", "DoLS", "COP"]
            }
        }

        for area_key, config in areas_config.items():
            # Area row
            area_container = QWidget()
            area_layout = QVBoxLayout(area_container)
            area_layout.setContentsMargins(0, 0, 0, 8)
            area_layout.setSpacing(4)

            # Label and dropdown row
            row = QHBoxLayout()
            lbl = QLabel(f"{config['label']}:")
            lbl.setFixedWidth(80)
            lbl.setStyleSheet("font-weight: 600; font-size: 15px;")
            row.addWidget(lbl)

            cap_combo = QComboBox()
            cap_combo.addItems(capacity_options)
            cap_combo.setStyleSheet(combo_style)
            row.addWidget(cap_combo)
            row.addStretch()
            area_layout.addLayout(row)

            # Actions container (hidden by default)
            actions_container = QWidget()
            actions_container.setVisible(False)

            action_widgets = {}

            if config["actions"] == "soad":
                # MHA paperwork and SOAD radios for medication - vertical layout
                actions_layout = QVBoxLayout(actions_container)
                actions_layout.setContentsMargins(80, 0, 0, 0)
                actions_layout.setSpacing(6)

                # First question: MHA paperwork in place?
                mha_row = QHBoxLayout()
                mha_row.setSpacing(8)
                mha_lbl = QLabel("MHA paperwork in place:")
                mha_lbl.setStyleSheet("font-size: 15px; color: #374151;")
                mha_row.addWidget(mha_lbl)

                mha_group = QButtonGroup(self)
                mha_yes = QRadioButton("Yes")
                mha_no = QRadioButton("No")
                mha_group.addButton(mha_yes)
                mha_group.addButton(mha_no)
                mha_row.addWidget(mha_yes)
                mha_row.addWidget(mha_no)
                mha_row.addStretch()
                actions_layout.addLayout(mha_row)

                # Second question: SOAD requested? (only visible if MHA = No)
                soad_container = QWidget()
                soad_container.setVisible(False)
                soad_row = QHBoxLayout(soad_container)
                soad_row.setContentsMargins(0, 0, 0, 0)
                soad_row.setSpacing(8)

                soad_lbl = QLabel("SOAD requested:")
                soad_lbl.setStyleSheet("font-size: 15px; color: #374151;")
                soad_row.addWidget(soad_lbl)

                soad_group = QButtonGroup(self)
                soad_yes = QRadioButton("Yes")
                soad_no = QRadioButton("No")
                soad_group.addButton(soad_yes)
                soad_group.addButton(soad_no)
                soad_row.addWidget(soad_yes)
                soad_row.addWidget(soad_no)
                soad_row.addStretch()
                actions_layout.addWidget(soad_container)

                # Connect MHA radios to show/hide SOAD question and update preview
                mha_yes.toggled.connect(lambda checked: soad_container.setVisible(not checked) if checked else None)
                mha_no.toggled.connect(lambda checked: soad_container.setVisible(checked))
                mha_yes.toggled.connect(lambda: self._update_preview("capacity"))
                mha_no.toggled.connect(lambda: self._update_preview("capacity"))
                soad_yes.toggled.connect(lambda: self._update_preview("capacity"))
                soad_no.toggled.connect(lambda: self._update_preview("capacity"))

                action_widgets["mha_yes"] = mha_yes
                action_widgets["mha_no"] = mha_no
                action_widgets["soad_yes"] = soad_yes
                action_widgets["soad_no"] = soad_no
                action_widgets["soad_container"] = soad_container
            elif config["actions"] == "finance_special":
                # Finance: IMCA checkbox + radio group for Guardianship/Appointeeship
                actions_layout = QVBoxLayout(actions_container)
                actions_layout.setContentsMargins(80, 0, 0, 0)
                actions_layout.setSpacing(6)

                # Checkboxes for Best Interest and IMCA
                cb_row = QHBoxLayout()
                cb_row.setSpacing(16)
                bi_cb = QCheckBox("Best Interest")
                bi_cb.setStyleSheet(cb_style)
                bi_cb.stateChanged.connect(lambda: self._update_preview("capacity"))
                cb_row.addWidget(bi_cb)
                imca_cb = QCheckBox("IMCA")
                imca_cb.setStyleSheet(cb_style)
                imca_cb.stateChanged.connect(lambda: self._update_preview("capacity"))
                cb_row.addWidget(imca_cb)
                cb_row.addStretch()
                actions_layout.addLayout(cb_row)

                # Radio buttons for Guardianship/Appointeeship (mutually exclusive)
                radio_row = QHBoxLayout()
                radio_row.setSpacing(12)
                fin_type_group = QButtonGroup(self)
                fin_none_rb = QRadioButton("None")
                fin_guardianship_rb = QRadioButton("Guardianship")
                fin_appointeeship_rb = QRadioButton("Appointeeship")
                fin_informal_rb = QRadioButton("Informal Appointeeship")
                fin_none_rb.setChecked(True)  # Default
                for rb in [fin_none_rb, fin_guardianship_rb, fin_appointeeship_rb, fin_informal_rb]:
                    rb.setStyleSheet("font-size: 15px; color: #374151;")
                    rb.toggled.connect(lambda: self._update_preview("capacity"))
                    fin_type_group.addButton(rb)
                    radio_row.addWidget(rb)
                radio_row.addStretch()
                actions_layout.addLayout(radio_row)

                action_widgets["best_interest"] = bi_cb
                action_widgets["imca"] = imca_cb
                action_widgets["fin_type_group"] = fin_type_group
                action_widgets["guardianship"] = fin_guardianship_rb
                action_widgets["appointeeship"] = fin_appointeeship_rb
                action_widgets["informal_appointeeship"] = fin_informal_rb
            else:
                # Checkboxes in 2-column grid layout
                actions_layout = QGridLayout(actions_container)
                actions_layout.setContentsMargins(80, 0, 0, 0)
                actions_layout.setSpacing(4)
                actions_layout.setHorizontalSpacing(16)

                for i, action in enumerate(config["actions"]):
                    cb = QCheckBox(action)
                    cb.setStyleSheet(cb_style)
                    cb.stateChanged.connect(lambda: self._update_preview("capacity"))
                    row = i // 2
                    col = i % 2
                    actions_layout.addWidget(cb, row, col)
                    action_widgets[action.lower().replace(" ", "_")] = cb

            area_layout.addWidget(actions_container)

            # Connect combo to show/hide actions
            cap_combo.currentIndexChanged.connect(
                lambda idx, cont=actions_container: cont.setVisible(idx == 2)  # "Lacks capacity"
            )
            cap_combo.currentIndexChanged.connect(lambda: self._update_preview("capacity"))

            self.popup_cap_areas[area_key] = {
                "combo": cap_combo,
                "actions_container": actions_container,
                "actions": action_widgets,
                "label": config["label"]
            }

            layout.addWidget(area_container)

        layout.addSpacing(8)

        # Additional notes
        layout.addWidget(QLabel("<b>Additional Details:</b>"))
        self.popup_cap_details = QTextEdit()
        self.popup_cap_details.setMaximumHeight(60)
        self.popup_cap_details.setPlaceholderText("Any additional capacity assessment details...")
        self.popup_cap_details.setStyleSheet("""
            QTextEdit { border: 1px solid #d1d5db; border-radius: 4px; padding: 8px; font-size: 16px; }
            QTextEdit:focus { border-color: #991b1b; }
        """)
        self.popup_cap_details.textChanged.connect(lambda: self._update_preview("capacity"))
        layout.addWidget(self.popup_cap_details)

        layout.addStretch()
        self._add_send_button(layout, "capacity", self._generate_capacity)

    def _generate_capacity(self) -> str:
        """Generate capacity text from per-area assessments and actions."""
        p = self._get_pronouns()
        parts = []

        has_capacity_areas = []
        lacks_capacity_parts = []

        for area_key, widgets in self.popup_cap_areas.items():
            combo_text = widgets["combo"].currentText()
            label = widgets["label"]

            if combo_text == "Has capacity":
                has_capacity_areas.append(label.lower())
            elif combo_text == "Lacks capacity":
                # Build sentence for lacking capacity
                area_part = f"{p['subj']} {p['lacks']} capacity for {label.lower()} decisions"

                # Get selected actions
                actions = widgets["actions"]
                selected_actions = []

                if area_key == "medication":
                    # Handle MHA paperwork and SOAD
                    if actions.get("mha_yes") and actions["mha_yes"].isChecked():
                        selected_actions.append("SOAD paperwork is in place")
                    elif actions.get("mha_no") and actions["mha_no"].isChecked():
                        if actions.get("soad_yes") and actions["soad_yes"].isChecked():
                            selected_actions.append("a SOAD has been requested")
                        elif actions.get("soad_no") and actions["soad_no"].isChecked():
                            selected_actions.append("a SOAD has not yet been requested")
                elif area_key == "finances":
                    # Handle finance special case - checkboxes + radio buttons
                    action_phrases = {
                        "best_interest": "a Best Interest decision is in place",
                        "imca": "an IMCA has been appointed",
                        "guardianship": "Guardianship is being considered",
                        "appointeeship": "an Appointeeship is in place",
                        "informal_appointeeship": "informal appointeeship arrangements are in place"
                    }
                    # Check the checkbox actions (Best Interest, IMCA)
                    for key in ["best_interest", "imca"]:
                        cb = actions.get(key)
                        if cb and hasattr(cb, 'isChecked') and cb.isChecked():
                            selected_actions.append(action_phrases[key])
                    # Check the radio button group (only one can be selected)
                    for key in ["guardianship", "appointeeship", "informal_appointeeship"]:
                        rb = actions.get(key)
                        if rb and hasattr(rb, 'isChecked') and rb.isChecked():
                            selected_actions.append(action_phrases[key])
                else:
                    # Handle checkboxes
                    action_phrases = {
                        "best_interest": "a Best Interest decision is in place",
                        "imca": "an IMCA has been appointed",
                        "dols": "a DoLS application has been made",
                        "cop": "a Court of Protection application is being considered",
                        "appointeeship": "an Appointeeship is in place",
                        "guardianship": "Guardianship is being considered",
                        "informal_appointeeship": "informal appointeeship arrangements are in place"
                    }
                    for action_key, cb in actions.items():
                        if hasattr(cb, 'isChecked') and cb.isChecked():
                            phrase = action_phrases.get(action_key, action_key)
                            selected_actions.append(phrase)

                if selected_actions:
                    area_part += " and " + ", ".join(selected_actions)

                lacks_capacity_parts.append(area_part + ".")

        # Build final output
        if has_capacity_areas:
            if len(has_capacity_areas) == 1:
                parts.append(f"{p['subj']} {p['has']} capacity for {has_capacity_areas[0]} decisions.")
            else:
                areas_text = ", ".join(has_capacity_areas[:-1]) + " and " + has_capacity_areas[-1]
                parts.append(f"{p['subj']} {p['has']} capacity for {areas_text} decisions.")

        if lacks_capacity_parts:
            parts.extend(lacks_capacity_parts)

        # Additional details
        details = self.popup_cap_details.toPlainText().strip()
        if details:
            parts.append(details)

        # Include checked imported entries
        if hasattr(self, 'popup_capacity_imported_entries'):
            for entry in self.popup_capacity_imported_entries:
                if entry.get("checkbox") and entry["checkbox"].isChecked():
                    full_text = entry.get("text", "").strip()
                    if full_text:
                        parts.append(f"\n\n[Imported Note]\n{full_text}")

        return " ".join(parts) if parts else "[Select capacity status for each area...]"

    def _build_popup_progress(self):
        """Build progress popup with sliders for each dimension and imported notes."""
        container, layout = self._create_popup_container_with_imports("progress")

        layout.addWidget(QLabel("Progress over the last 12 months:"))
        layout.addSpacing(10)

        # Slider style
        slider_style = """
            QSlider::groove:horizontal {
                height: 6px;
                background: #e5e7eb;
                border-radius: 3px;
            }
            QSlider::handle:horizontal {
                width: 18px;
                margin: -6px 0;
                background: #991b1b;
                border-radius: 9px;
            }
            QSlider::sub-page:horizontal {
                background: #991b1b;
                border-radius: 3px;
            }
        """

        # === MENTAL STATE ===
        self.progress_mental_options = [
            "Unsettled", "Often unsettled", "Unsettled at times", "Stable",
            "Some improvement", "Significant improvement", "Symptom free with no concerns"
        ]
        layout.addWidget(QLabel("<b>Mental State:</b>"))
        self.progress_mental_label = QLabel(self.progress_mental_options[3])
        self.progress_mental_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        layout.addWidget(self.progress_mental_label)
        self.progress_mental_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_mental_slider.setRange(0, len(self.progress_mental_options) - 1)
        self.progress_mental_slider.setValue(3)
        self.progress_mental_slider.setStyleSheet(slider_style)
        self.progress_mental_slider.valueChanged.connect(
            lambda v: self.progress_mental_label.setText(self.progress_mental_options[v])
        )
        layout.addWidget(self.progress_mental_slider)
        layout.addSpacing(8)

        # === INSIGHT ===
        self.progress_insight_options = [
            "Remains limited", "Mostly absent but some insight", "Mild insight",
            "Moderate insight", "Good insight", "Full insight"
        ]
        layout.addWidget(QLabel("<b>Insight:</b>"))
        self.progress_insight_label = QLabel(self.progress_insight_options[2])
        self.progress_insight_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        layout.addWidget(self.progress_insight_label)
        self.progress_insight_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_insight_slider.setRange(0, len(self.progress_insight_options) - 1)
        self.progress_insight_slider.setValue(2)
        self.progress_insight_slider.setStyleSheet(slider_style)
        self.progress_insight_slider.valueChanged.connect(
            lambda v: self.progress_insight_label.setText(self.progress_insight_options[v])
        )
        layout.addWidget(self.progress_insight_slider)
        layout.addSpacing(8)

        # === ENGAGEMENT ===
        self.progress_engagement_options = ["Nil", "Some", "Partial", "Good", "Very good", "Full"]
        layout.addWidget(QLabel("<b>Engagement with Treatment:</b>"))
        self.progress_engagement_label = QLabel(self.progress_engagement_options[3])
        self.progress_engagement_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        layout.addWidget(self.progress_engagement_label)
        self.progress_engagement_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_engagement_slider.setRange(0, len(self.progress_engagement_options) - 1)
        self.progress_engagement_slider.setValue(3)
        self.progress_engagement_slider.setStyleSheet(slider_style)
        self.progress_engagement_slider.valueChanged.connect(
            lambda v: self.progress_engagement_label.setText(self.progress_engagement_options[v])
        )
        layout.addWidget(self.progress_engagement_slider)
        layout.addSpacing(8)

        # === RISK REDUCTION WORK ===
        self.progress_risk_options = ["Nil", "Started", "In process", "Good engagement", "Almost completed", "Completed"]
        layout.addWidget(QLabel("<b>Risk Reduction Work:</b>"))
        self.progress_risk_label = QLabel(self.progress_risk_options[2])
        self.progress_risk_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        layout.addWidget(self.progress_risk_label)
        self.progress_risk_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_risk_slider.setRange(0, len(self.progress_risk_options) - 1)
        self.progress_risk_slider.setValue(2)
        self.progress_risk_slider.setStyleSheet(slider_style)
        self.progress_risk_slider.valueChanged.connect(
            lambda v: self.progress_risk_label.setText(self.progress_risk_options[v])
        )
        layout.addWidget(self.progress_risk_slider)
        layout.addSpacing(8)

        # === LEAVE TYPE ===
        self.progress_leave_options = ["No leave", "Escorted", "Unescorted", "Overnight"]
        layout.addWidget(QLabel("<b>Leave:</b>"))
        self.progress_leave_label = QLabel(self.progress_leave_options[0])
        self.progress_leave_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        layout.addWidget(self.progress_leave_label)
        self.progress_leave_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_leave_slider.setRange(0, len(self.progress_leave_options) - 1)
        self.progress_leave_slider.setValue(0)
        self.progress_leave_slider.setStyleSheet(slider_style)
        self.progress_leave_slider.valueChanged.connect(self._on_progress_leave_changed)
        layout.addWidget(self.progress_leave_slider)
        layout.addSpacing(8)

        # === LEAVE USAGE (conditional) ===
        self.progress_leave_usage_container = QWidget()
        leave_usage_layout = QVBoxLayout(self.progress_leave_usage_container)
        leave_usage_layout.setContentsMargins(0, 0, 0, 0)
        leave_usage_layout.setSpacing(4)
        self.progress_usage_options = ["Intermittent", "Variable", "Regular", "Frequent", "Excellent"]
        leave_usage_layout.addWidget(QLabel("<b>Leave Usage:</b>"))
        self.progress_usage_label = QLabel(self.progress_usage_options[2])
        self.progress_usage_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        leave_usage_layout.addWidget(self.progress_usage_label)
        self.progress_usage_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_usage_slider.setRange(0, len(self.progress_usage_options) - 1)
        self.progress_usage_slider.setValue(2)
        self.progress_usage_slider.setStyleSheet(slider_style)
        self.progress_usage_slider.valueChanged.connect(
            lambda v: self.progress_usage_label.setText(self.progress_usage_options[v])
        )
        leave_usage_layout.addWidget(self.progress_usage_slider)
        self.progress_leave_usage_container.hide()
        layout.addWidget(self.progress_leave_usage_container)

        # === LEAVE CONCERNS (conditional) ===
        self.progress_leave_concerns_container = QWidget()
        leave_concerns_layout = QVBoxLayout(self.progress_leave_concerns_container)
        leave_concerns_layout.setContentsMargins(0, 0, 0, 0)
        leave_concerns_layout.setSpacing(4)
        self.progress_concerns_options = ["No", "Mild", "Some", "Several", "Significant"]
        leave_concerns_layout.addWidget(QLabel("<b>Leave Concerns:</b>"))
        self.progress_concerns_label = QLabel(self.progress_concerns_options[0])
        self.progress_concerns_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        leave_concerns_layout.addWidget(self.progress_concerns_label)
        self.progress_concerns_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_concerns_slider.setRange(0, len(self.progress_concerns_options) - 1)
        self.progress_concerns_slider.setValue(0)
        self.progress_concerns_slider.setStyleSheet(slider_style)
        self.progress_concerns_slider.valueChanged.connect(
            lambda v: self.progress_concerns_label.setText(self.progress_concerns_options[v])
        )
        leave_concerns_layout.addWidget(self.progress_concerns_slider)
        self.progress_leave_concerns_container.hide()
        layout.addWidget(self.progress_leave_concerns_container)
        layout.addSpacing(8)

        # === DISCHARGE PLANNING ===
        self.progress_discharge_options = ["Not started", "Early stages", "In progress", "Almost completed", "Completed"]
        layout.addWidget(QLabel("<b>Discharge Planning:</b>"))
        self.progress_discharge_label = QLabel(self.progress_discharge_options[0])
        self.progress_discharge_label.setStyleSheet("color: #991b1b; font-weight: 600;")
        layout.addWidget(self.progress_discharge_label)
        self.progress_discharge_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.progress_discharge_slider.setRange(0, len(self.progress_discharge_options) - 1)
        self.progress_discharge_slider.setValue(0)
        self.progress_discharge_slider.setStyleSheet(slider_style)
        self.progress_discharge_slider.valueChanged.connect(
            lambda v: self.progress_discharge_label.setText(self.progress_discharge_options[v])
        )
        layout.addWidget(self.progress_discharge_slider)
        layout.addSpacing(8)

        # === DISCHARGE APPLICATIONS ===
        layout.addWidget(QLabel("<b>Discharge Applications:</b>"))
        app_row = QHBoxLayout()
        self.progress_app_group = QButtonGroup(self)
        self.progress_app_yes = QRadioButton("Yes")
        self.progress_app_no = QRadioButton("No")
        self.progress_app_no.setChecked(True)
        self.progress_app_group.addButton(self.progress_app_yes)
        self.progress_app_group.addButton(self.progress_app_no)
        app_row.addWidget(self.progress_app_yes)
        app_row.addWidget(self.progress_app_no)
        app_row.addStretch()
        layout.addLayout(app_row)

        layout.addStretch()
        self._add_send_button(layout, "progress", self._generate_progress)
        self._connect_preview_updates("progress", [
            self.progress_mental_slider, self.progress_insight_slider,
            self.progress_engagement_slider, self.progress_risk_slider,
            self.progress_leave_slider, self.progress_discharge_slider,
            self.progress_app_yes, self.progress_app_no
        ])

        # Initialize tracking for Section 8 narrative display (matching Tribunal Section 14 style)
        self._progress_narrative_text = ""
        self._progress_narrative_html = ""
        self._progress_entry_frames = {}
        self._progress_entry_body_texts = {}
        self._progress_extracted_checkboxes = []

    def _on_progress_leave_changed(self, value):
        """Show/hide leave usage and concerns based on leave type."""
        self.progress_leave_label.setText(self.progress_leave_options[value])
        has_leave = value > 0  # Not "No leave"
        self.progress_leave_usage_container.setVisible(has_leave)
        self.progress_leave_concerns_container.setVisible(has_leave)

    def _generate_progress(self) -> str:
        """Generate narrative progress text."""
        p = self._get_pronouns()

        # Get values
        mental = self.progress_mental_options[self.progress_mental_slider.value()].lower()
        insight = self.progress_insight_options[self.progress_insight_slider.value()].lower()
        engagement = self.progress_engagement_options[self.progress_engagement_slider.value()].lower()
        risk_work = self.progress_risk_options[self.progress_risk_slider.value()].lower()
        leave_type = self.progress_leave_options[self.progress_leave_slider.value()].lower()
        discharge = self.progress_discharge_options[self.progress_discharge_slider.value()].lower()
        has_application = self.progress_app_yes.isChecked()

        # Build narrative
        parts = []
        parts.append(f"Over the last 12 months, mental state has been {mental}.")

        # Insight phrase
        if "full" in insight:
            parts.append(f"{p['subj']} {p['has']} displayed full insight into {p['pos_l']} needs and {p['pos_l']} illness.")
        elif "good" in insight:
            parts.append(f"{p['subj']} {p['has']} displayed good insight into {p['pos_l']} needs and {p['pos_l']} illness.")
        elif "moderate" in insight:
            parts.append(f"{p['subj']} {p['has']} displayed moderate insight into {p['pos_l']} needs and {p['pos_l']} illness.")
        elif "mild" in insight:
            parts.append(f"{p['subj']} {p['has']} displayed some mild insight into {p['pos_l']} needs and {p['pos_l']} illness.")
        elif "mostly absent" in insight:
            parts.append(f"{p['pos']} insight remains mostly absent though there are some signs of emerging awareness.")
        else:
            parts.append(f"{p['pos']} insight remains limited.")

        # Engagement
        parts.append(f"{p['pos']} engagement with treatment has been {engagement} overall")

        # Leave section
        if leave_type == "no leave":
            parts[-1] += f" and {p['subj_l']} {p['has']} not been taking any leave."
        else:
            usage = self.progress_usage_options[self.progress_usage_slider.value()].lower()
            concerns = self.progress_concerns_options[self.progress_concerns_slider.value()].lower()
            leave_phrase = f"and {p['subj_l']} {p['has']} been taking {leave_type} leave on a {usage} basis"
            if concerns == "no":
                leave_phrase += " without concerns."
            else:
                leave_phrase += f" with {concerns} concerns."
            parts[-1] += " " + leave_phrase

        # Discharge planning
        if discharge == "not started":
            parts.append("Currently discharge planning has not started.")
        elif discharge == "early stages":
            parts.append("Currently discharge planning is at an early stage.")
        elif discharge == "in progress":
            parts.append("Discharge planning is currently in progress.")
        elif discharge == "almost completed":
            parts.append("Discharge planning is almost completed.")
        else:
            parts.append("Discharge planning has been completed.")

        # Applications
        if has_application:
            parts.append("During this period applications have been made for discharge.")
        else:
            parts.append("During this period no applications have been made for discharge.")

        return " ".join(parts)

    def populate_popup_progress_imports(self, entries: list):
        """Populate the imported data panel in popup 8 (Progress) matching Tribunal Section 14 style.

        Generates a narrative summary with clickable references and collapsible entry boxes.
        Filters to 12-month window from latest entry.
        """
        from datetime import datetime, timedelta
        from collections import defaultdict
        from pathlib import Path
        import re
        import html as html_module
        from progress_panel import reset_reference_tracker, make_link, get_reference

        print(f"[MOJ-ASR] Section 8 popup: populate_popup_progress_imports called with {len(entries) if entries else 0} entries")

        import_layout = getattr(self, 'popup_progress_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 8 popup: popup_progress_import_layout not available")
            return

        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_progress_imported_entries = []
        self._progress_entry_frames = {}
        self._progress_entry_body_texts = {}
        self._progress_extracted_checkboxes = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Reset reference tracker for fresh narrative
        reset_reference_tracker()

        # Get patient info for narrative generation
        from shared_data_store import get_shared_store
        store = get_shared_store()
        patient_info = store.patient_info
        patient_name = patient_info.get("name", "The patient")
        name = patient_name.split()[0] if patient_name else "The patient"

        # Set up pronouns using central utilities
        pronouns = store.gender_pronouns  # Uses patient_demographics.get_pronouns()
        pronoun = pronouns['subject']
        pronoun_obj = pronouns['object']
        pronoun_poss = pronouns['possessive']
        pronoun_cap = pronoun.capitalize()
        pronoun_poss_cap = pronoun_poss.capitalize()

        # Progress/Mental State categories based on tribunal section 14 keywords
        # Note: Use word boundary matching (\\b) for keywords that could be substrings of other words
        CATEGORIES = {
            "Mental State": [
                "mental state", "mse", "mental status", "presentation",
                "settled", "unsettled", "stable", "unstable", "calm", "agitated",
                "irritable", "anxious", "low mood", "elated", "elevated mood",
                "psychotic", "paranoid", "guarded", "suspicious", "thought disorder",
                "hallucination", "delusion", "voices", "hearing voices"
            ],
            "Positive Progress": [
                "progress", "improvement", "improved", "improving", "better",
                "settled in his mental", "settled in her mental", "bright in mood",
                "\\bappropriate\\b", "pleasant", "cooperative", "engaging well",
                "good rapport", "well presented", "\\bstable\\b", "no concerns",
                "no challenging behaviour", "no incidents"
            ],
            "Deterioration": [
                "deteriorat", "decline", "worsening", "relapse", "decompensate",
                "acutely unwell", "becoming unwell", "symptoms returning",
                "more unwell", "less stable", "increased symptoms",
                "inappropriate", "\\bunstable\\b"
            ],
            "Insight": [
                "insight", "awareness", "understanding", "recogni", "acknowledge",
                "accept", "denial", "denies", "lack of insight", "poor insight",
                "good insight", "partial insight", "limited insight", "believes",
                "doesn't believe", "mental illness", "illness awareness"
            ],
            "Engagement": [
                "engag", "disengag", "attend", "did not attend", "dna",
                "rapport", "therapeutic relationship", "working with", "cooperat",
                "uncooperative", "resistant", "reluctant", "willing", "unwilling",
                "participat", "involved", "motivation", "motivated"
            ],
            "Risk Work": [
                "psychology", "psychologist", "cbt", "dbt", "violence reduction",
                "offence related work", "offending behaviour", "formulation",
                "therapy", "treatment programme", "risk reduction", "ot group",
                "occupational therapy", "meaningful activity"
            ],
            "Leave": [
                "leave", "escorted", "unescorted", "overnight", "ground leave",
                "community leave", "section 17", "s17", "trial leave", "home leave"
            ],
            "Medication": [
                "medication", "clozapine", "depot", "antipsychotic", "olanzapine",
                "risperidone", "aripiprazole", "quetiapine", "haloperidol",
                "concordant", "compliant with medication", "taking medication",
                "refused medication", "declined medication"
            ],
        }

        CATEGORY_COLORS = {
            "Mental State": "#7c3aed",      # Purple
            "Positive Progress": "#059669", # Green
            "Deterioration": "#dc2626",     # Red
            "Insight": "#0891b2",           # Cyan
            "Engagement": "#2563eb",        # Blue
            "Risk Work": "#d97706",         # Amber
            "Leave": "#0d9488",             # Teal
            "Medication": "#be185d",        # Pink
        }

        self._progress_categories = CATEGORIES
        self._progress_colors = CATEGORY_COLORS
        self._progress_current_filter = None

        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def keyword_matches(keyword, text_lower):
            """Check if keyword matches text, using regex for word boundary patterns."""
            if '\\b' in keyword:
                # Use regex matching for word boundary patterns
                pattern = keyword.replace('\\\\b', '\\b')  # Handle escaped backslashes
                return re.search(pattern, text_lower, re.IGNORECASE) is not None
            else:
                # Simple substring matching
                return keyword in text_lower

        def extract_relevant_snippet(text, matched_categories):
            lines = text.strip().split('\n')
            context = '\n'.join(lines[:2]).strip()
            relevant_sentences = []
            sentences = re.split(r'(?<=[.!?])\s+', text)

            for cat in matched_categories:
                keywords = CATEGORIES.get(cat, [])
                for kw in keywords:
                    for sent in sentences:
                        if keyword_matches(kw, sent.lower()) and sent.strip() not in relevant_sentences:
                            if sent.strip() not in context:
                                relevant_sentences.append(sent.strip())
                            break

            snippet = context
            if relevant_sentences:
                additional = ' ... '.join(relevant_sentences[:2])
                if additional and additional not in context:
                    snippet += f"\n[...] {additional}"
            return snippet if snippet else text[:200]

        def categorize_text(text):
            text_lower = text.lower()
            matches = []
            for cat, keywords in CATEGORIES.items():
                for kw in keywords:
                    if keyword_matches(kw, text_lower):
                        matches.append(cat)
                        break
            return matches

        categorized = []
        for entry in filtered_entries:
            cats = categorize_text(entry["text"])
            if cats:
                snippet = extract_relevant_snippet(entry["text"], cats)
                entry["snippet"] = snippet
                categorized.append((entry, cats))

        categorized.sort(key=lambda x: x[0].get("date_obj") or datetime.min, reverse=True)

        # Add category labels to header showing what categories were found (clickable to filter)
        category_buttons_layout = getattr(self, 'popup_progress_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            category_counts = {}
            for entry, cats in categorized:
                for cat in cats:
                    category_counts[cat] = category_counts.get(cat, 0) + 1
            if category_counts:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
            for cat in sorted(category_counts.keys()):
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                count = category_counts[cat]
                btn = QPushButton(f"{cat} ({count})")
                btn.setCursor(Qt.CursorShape.PointingHandCursor)
                btn.setStyleSheet(f"QPushButton {{ font-size: 14px; font-weight: 600; color: white; background: {color}; padding: 2px 6px; border-radius: 4px; border: none; }} QPushButton:hover {{ opacity: 0.8; }}")
                btn.clicked.connect(lambda checked, c=cat: self._apply_progress_filter(c))
                category_buttons_layout.addWidget(btn)

        if not categorized:
            placeholder = QLabel("No progress-related notes found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        self._progress_all_categorized = categorized

        # ========== GENERATE NARRATIVE SUMMARY ==========
        # Convert categorized entries to format for narrative generation
        entries_for_narrative = []
        for entry, cats in categorized:
            entries_for_narrative.append({
                'date': entry.get('date_obj'),
                'text': entry.get('text', ''),
                'content': entry.get('text', ''),
                'categories': cats
            })

        # Generate narrative summary using the SAME comprehensive function as Tribunal Section 14
        # Import the tribunal's TribunalProgressPopup to use its comprehensive narrative generator
        # IMPORTANT: Use ALL filtered_entries, not just categorized ones - the tribunal function
        # analyzes ALL entries for professional contacts, risk scoring, themes, etc.
        try:
            from tribunal_popups import TribunalProgressPopup

            # Convert ALL filtered_entries (not just categorized) to format expected by tribunal
            # Tribunal expects: content or text, and date or datetime
            tribunal_entries = []
            for entry in filtered_entries:  # Use filtered_entries, NOT entries_for_narrative
                tribunal_entries.append({
                    'content': entry.get('text', ''),
                    'text': entry.get('text', ''),
                    'date': entry.get('date_obj'),  # Use date_obj which is the parsed datetime
                    'datetime': entry.get('date_obj'),
                    'type': entry.get('type', 'Progress Note'),
                })

            # Create temporary instance to call the comprehensive narrative generator
            temp_popup = TribunalProgressPopup.__new__(TribunalProgressPopup)
            plain_text, narrative_html = temp_popup._generate_narrative_summary(tribunal_entries)
            self._progress_narrative_text = plain_text
            self._progress_narrative_html = narrative_html
            print(f"[MOJ-ASR] Section 8: Generated comprehensive narrative using tribunal function ({len(tribunal_entries)} entries)")
        except Exception as e:
            print(f"[MOJ-ASR] Section 8: Failed to use tribunal narrative, falling back to simple: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to simple narrative
            narrative_html = self._generate_progress_narrative(entries_for_narrative, name, pronoun, pronoun_cap, pronoun_poss, pronoun_poss_cap)
            self._progress_narrative_html = narrative_html

        # ========== CREATE NARRATIVE SECTION ==========
        # Use CollapsibleSection if available (imported at top of file)

        # Clinical Narrative Summary section
        if CollapsibleSection:
            narrative_section = CollapsibleSection("Clinical Narrative Summary")
        else:
            narrative_section = QFrame()
            narrative_section.setObjectName("narrative_section")
        narrative_section.setStyleSheet("""
            CollapsibleSection {
                background: rgba(255, 248, 220, 0.8);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 8px;
            }
        """)

        narrative_content = QWidget()
        narrative_content.setStyleSheet("background: transparent;")
        narrative_layout = QVBoxLayout(narrative_content)
        narrative_layout.setContentsMargins(8, 8, 8, 8)
        narrative_layout.setSpacing(8)

        # Include narrative checkbox
        include_narrative_cb = QCheckBox("Include narrative in output")
        include_narrative_cb.setChecked(True)
        include_narrative_cb.setStyleSheet("QCheckBox { font-size: 15px; color: #806000; background: transparent; }")
        include_narrative_cb.stateChanged.connect(lambda state: self._update_preview("progress"))
        narrative_layout.addWidget(include_narrative_cb)
        self._progress_include_narrative_cb = include_narrative_cb

        # Narrative text browser
        narrative_text_widget = QTextBrowser()
        narrative_text_widget.setOpenExternalLinks(False)
        narrative_text_widget.setOpenLinks(False)
        narrative_text_widget.anchorClicked.connect(self._on_progress_narrative_link_clicked)
        narrative_text_widget.setStyleSheet("""
            QTextBrowser {
                font-size: 16px;
                color: #333;
                background: rgba(255, 255, 255, 0.7);
                border: none;
                border-radius: 4px;
                padding: 8px;
            }
        """)
        narrative_text_widget.setHtml(narrative_html)
        narrative_text_widget.setMinimumHeight(150)
        narrative_layout.addWidget(narrative_text_widget)
        self._progress_narrative_text_widget = narrative_text_widget

        if CollapsibleSection:
            narrative_section.set_content(narrative_content)
        else:
            narrative_section_layout = QVBoxLayout(narrative_section)
            narrative_section_layout.addWidget(narrative_content)
        import_layout.addWidget(narrative_section)

        # ========== CREATE ENTRIES SECTION ==========
        if CollapsibleSection:
            entries_section = CollapsibleSection("Individual Progress Notes")
            entries_section.setStyleSheet("""
                CollapsibleSection {
                    background: rgba(255, 248, 220, 0.6);
                    border: 1px solid rgba(180, 150, 50, 0.3);
                    border-radius: 8px;
                }
            """)
        else:
            entries_section = QFrame()
            entries_section.setObjectName("entries_section")
            entries_section.setStyleSheet("""
                QFrame#entries_section {
                    background: rgba(255, 248, 220, 0.6);
                    border: 1px solid rgba(180, 150, 50, 0.3);
                    border-radius: 8px;
                }
            """)

        entries_content = QWidget()
        entries_content.setStyleSheet("background: transparent;")
        entries_layout = QVBoxLayout(entries_content)
        entries_layout.setContentsMargins(4, 4, 4, 4)
        entries_layout.setSpacing(8)
        entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                # Handle word boundary patterns vs regular keywords
                if '\\b' in kw:
                    # Use the pattern directly (convert escaped backslashes)
                    pattern_str = kw.replace('\\\\b', '\\b')
                    pattern = re.compile(pattern_str, re.IGNORECASE)
                else:
                    pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        def create_entry_widget(entry, categories, filter_callback):
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_str = format_date_nice(date_obj) if date_obj else ""
            highlighted_html = highlight_keywords(full_text, categories or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="progress",
                categories=categories,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=filter_callback,
                parent_container=getattr(self, 'popup_progress_import_container', None),
            )
            return entry_frame, cb, full_text

        imported_entries = []
        for idx, (entry, cats) in enumerate(categorized):
            widget, cb, full_text = create_entry_widget(entry, cats, self._apply_progress_filter)
            entries_layout.addWidget(widget)
            imported_entries.append({"checkbox": cb, "text": full_text, "date": entry.get("date", ""), "categories": cats})
            self._progress_extracted_checkboxes.append(cb)

            # Store references for scrolling from narrative links
            date_obj = entry.get("date_obj")
            if date_obj:
                date_key = date_obj.strftime("%d/%m/%Y")
                unique_key = f"{date_key}_{idx}"
                self._progress_entry_frames[unique_key] = widget
                self._progress_entry_frames[date_key] = widget  # Also store by date

        # Finalize entries section
        if CollapsibleSection:
            entries_section.set_content(entries_content)
        else:
            entries_section_layout = QVBoxLayout(entries_section)
            entries_section_layout.addWidget(entries_content)
        import_layout.addWidget(entries_section)

        self.popup_progress_imported_entries = imported_entries
        print(f"[MOJ-ASR] Section 8 popup: Displayed {len(categorized)} categorized progress entries")

    def _generate_progress_narrative(self, entries, name, pronoun, pronoun_cap, pronoun_poss, pronoun_poss_cap):
        """Generate narrative summary for Section 8 Progress (matching Tribunal Section 14 style)."""
        from datetime import datetime, timedelta
        from collections import defaultdict
        from progress_panel import make_link
        import re

        if not entries:
            return ""

        # Sort entries by date
        sorted_entries = sorted([e for e in entries if e.get('date')], key=lambda x: x['date'])
        if not sorted_entries:
            return ""

        earliest = sorted_entries[0]['date']
        latest = sorted_entries[-1]['date']
        date_range = (latest - earliest).days
        date_range_months = max(1, date_range // 30)

        narrative_parts = []

        # === OVERVIEW ===
        narrative_parts.append("<b>REVIEW PERIOD OVERVIEW</b>")

        total_entries = len(sorted_entries)
        freq_desc = "regularly" if total_entries > date_range_months * 2 else "periodically"
        narrative_parts.append(f"{name} has been reviewed {freq_desc} over the past {date_range_months} months with {total_entries} documented contacts.")

        # Count by category
        category_counts = defaultdict(int)
        for entry in entries:
            for cat in entry.get('categories', []):
                category_counts[cat] += 1

        if category_counts:
            top_categories = sorted(category_counts.items(), key=lambda x: -x[1])[:3]
            cat_summary = ", ".join([f"{cat} ({count})" for cat, count in top_categories])
            narrative_parts.append(f"Key themes documented: {cat_summary}.")

        narrative_parts.append("")

        # === MENTAL STATE SUMMARY ===
        mental_state_entries = [e for e in entries if 'Mental State' in e.get('categories', [])]
        positive_entries = [e for e in entries if 'Positive Progress' in e.get('categories', [])]
        deterioration_entries = [e for e in entries if 'Deterioration' in e.get('categories', [])]

        narrative_parts.append("<b>MENTAL STATE AND PROGRESS</b>")

        if len(positive_entries) > len(deterioration_entries) * 2:
            ms_summary = f"{name}'s mental state has been predominantly stable and positive throughout the review period."
        elif len(deterioration_entries) > len(positive_entries):
            ms_summary = f"{name}'s mental state has shown some variability with periods of concern during the review period."
        else:
            ms_summary = f"{name}'s mental state has been variable, with both positive presentations and some concerning periods documented."

        narrative_parts.append(ms_summary)

        # Most recent positive
        if positive_entries:
            recent_pos = sorted(positive_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_pos.get('text', '')[:100]
            link = make_link(recent_pos['date'].strftime('%d %B %Y'), recent_pos['date'], 'positive', excerpt)
            narrative_parts.append(f"Most recently, on {link}, {pronoun} presented positively.")

        # Most recent concern
        if deterioration_entries:
            recent_neg = sorted(deterioration_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_neg.get('text', '')[:100]
            link = make_link(recent_neg['date'].strftime('%d %B %Y'), recent_neg['date'], 'deterioration', excerpt)
            narrative_parts.append(f"Concerns were noted on {link}.")

        narrative_parts.append("")

        # === ENGAGEMENT ===
        engagement_entries = [e for e in entries if 'Engagement' in e.get('categories', [])]
        if engagement_entries:
            narrative_parts.append("<b>ENGAGEMENT WITH SERVICES</b>")

            eng_count = len(engagement_entries)
            link = make_link(f"{eng_count} entries", None, 'engagement', '')
            narrative_parts.append(f"Engagement with services was documented in {link}.")

            recent_eng = sorted(engagement_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_eng.get('text', '')[:100]
            eng_link = make_link(recent_eng['date'].strftime('%d %B'), recent_eng['date'], 'engagement', excerpt)
            narrative_parts.append(f"Most recent engagement note on {eng_link}.")

            narrative_parts.append("")

        # === INSIGHT ===
        insight_entries = [e for e in entries if 'Insight' in e.get('categories', [])]
        if insight_entries:
            narrative_parts.append("<b>INSIGHT</b>")

            ins_count = len(insight_entries)
            link = make_link(f"{ins_count} entries", None, 'insight', '')
            narrative_parts.append(f"Insight was documented in {link}.")

            recent_ins = sorted(insight_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_ins.get('text', '')[:100]
            ins_link = make_link(recent_ins['date'].strftime('%d %B'), recent_ins['date'], 'insight', excerpt)
            narrative_parts.append(f"Most recent insight documentation on {ins_link}.")

            narrative_parts.append("")

        # === LEAVE ===
        leave_entries = [e for e in entries if 'Leave' in e.get('categories', [])]
        if leave_entries:
            narrative_parts.append("<b>LEAVE</b>")

            leave_count = len(leave_entries)
            link = make_link(f"{leave_count} entries", None, 'leave', '')
            narrative_parts.append(f"Leave was documented in {link}.")

            recent_leave = sorted(leave_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_leave.get('text', '')[:100]
            leave_link = make_link(recent_leave['date'].strftime('%d %B'), recent_leave['date'], 'leave', excerpt)
            narrative_parts.append(f"Most recent leave entry on {leave_link}.")

            narrative_parts.append("")

        # === MEDICATION ===
        medication_entries = [e for e in entries if 'Medication' in e.get('categories', [])]
        if medication_entries:
            narrative_parts.append("<b>MEDICATION</b>")

            med_count = len(medication_entries)
            link = make_link(f"{med_count} entries", None, 'medication', '')
            narrative_parts.append(f"Medication was documented in {link}.")

            recent_med = sorted(medication_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_med.get('text', '')[:100]
            med_link = make_link(recent_med['date'].strftime('%d %B'), recent_med['date'], 'medication', excerpt)
            narrative_parts.append(f"Most recent medication review on {med_link}.")

            narrative_parts.append("")

        # === RISK WORK ===
        risk_work_entries = [e for e in entries if 'Risk Work' in e.get('categories', [])]
        if risk_work_entries:
            narrative_parts.append("<b>RISK REDUCTION WORK</b>")

            rw_count = len(risk_work_entries)
            link = make_link(f"{rw_count} entries", None, 'risk work', '')
            narrative_parts.append(f"Risk reduction work was documented in {link}.")

            recent_rw = sorted(risk_work_entries, key=lambda x: x['date'])[-1]
            excerpt = recent_rw.get('text', '')[:100]
            rw_link = make_link(recent_rw['date'].strftime('%d %B'), recent_rw['date'], 'risk work', excerpt)
            narrative_parts.append(f"Most recent entry on {rw_link}.")

            narrative_parts.append("")

        # === SUMMARY ===
        narrative_parts.append("<b>SUMMARY</b>")
        summary_points = []
        summary_points.append(f"{name} has been reviewed {freq_desc} over the past {date_range_months} months")

        if len(positive_entries) > len(deterioration_entries):
            summary_points.append(f"{pronoun_poss} mental state has been predominantly stable")
        else:
            summary_points.append(f"{pronoun_poss} mental state has shown some variability")

        if engagement_entries:
            summary_points.append(f"with documented engagement with services")

        narrative_parts.append(". ".join(summary_points) + ".")

        # Format as HTML
        html_text = "<br>".join(narrative_parts)

        styled_html = f"""
        <style>
            a {{ color: #0066cc; text-decoration: underline; cursor: pointer; }}
            a:hover {{ background-color: rgba(255, 200, 0, 0.3); }}
        </style>
        <div style='font-family: sans-serif; font-size: 15px; color: #1f2937; line-height: 1.5;'>{html_text}</div>
        """

        return styled_html

    def _on_progress_narrative_link_clicked(self, url):
        """Handle clicks on narrative reference links in Section 8 Progress popup."""
        from PySide6.QtCore import QTimer
        from progress_panel import get_reference
        import re

        ref_id = url.fragment()
        if not ref_id:
            return

        ref_data = get_reference(ref_id)
        if not ref_data:
            return

        matched_text = ref_data.get("matched", "")
        ref_date = ref_data.get("date")

        # Find matching entries and scroll to them
        if not hasattr(self, '_progress_entry_frames'):
            return

        # Try to find entry by date
        if ref_date:
            if hasattr(ref_date, 'strftime'):
                date_key = ref_date.strftime("%d/%m/%Y")
            else:
                date_key = str(ref_date)

            entry_frame = self._progress_entry_frames.get(date_key)
            if entry_frame:
                # Scroll to entry
                QTimer.singleShot(100, lambda: entry_frame.ensurePolished())
                entry_frame.setStyleSheet("QFrame { background: #fffbeb; border: 2px solid #f59e0b; border-radius: 4px; }")
                QTimer.singleShot(2000, lambda: entry_frame.setStyleSheet("QFrame { background: white; border: none; border-radius: 4px; }"))

    def _apply_progress_filter(self, category: str):
        """Filter progress entries by category."""
        if not hasattr(self, '_progress_all_categorized'):
            return

        import_layout = getattr(self, 'popup_progress_import_layout', None)
        if not import_layout:
            return

        CATEGORIES = self._progress_categories
        CATEGORY_COLORS = self._progress_colors

        # Toggle filter
        if self._progress_current_filter == category:
            self._progress_current_filter = None
        else:
            self._progress_current_filter = category

        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Add filter bar when filter is active (fixed height for consistency)
        if self._progress_current_filter:
            filter_frame = QFrame()
            filter_frame.setFixedHeight(28)
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 0)
            filter_row.setSpacing(8)

            color = CATEGORY_COLORS.get(self._progress_current_filter, "#6b7280")
            filter_label = QLabel(f"Filtered by: {self._progress_current_filter}")
            filter_label.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {color}; background: transparent;")
            filter_row.addWidget(filter_label)
            filter_row.addStretch()

            remove_btn = QPushButton("Remove filter")
            remove_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            remove_btn.setFixedHeight(22)
            remove_btn.setStyleSheet("""
                QPushButton {
                    font-size: 14px; color: #dc2626; background: transparent;
                    border: 1px solid #dc2626; border-radius: 4px; padding: 2px 8px;
                }
                QPushButton:hover { background: #fef2f2; }
            """)
            remove_btn.clicked.connect(lambda: self._apply_progress_filter(self._progress_current_filter))
            filter_row.addWidget(remove_btn)

            import_layout.addWidget(filter_frame)

        filtered = []
        for entry, cats in self._progress_all_categorized:
            if self._progress_current_filter is None or self._progress_current_filter in cats:
                filtered.append((entry, cats))

        if not filtered:
            placeholder = QLabel(f"No entries match filter: {self._progress_current_filter}")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_categories):
            import html
            import re
            escaped = html.escape(text)
            keyword_colors = {}
            for cat in matched_categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                light_color = hex_to_light_bg(color)
                for kw in CATEGORIES.get(cat, []):
                    keyword_colors[kw] = light_color
            sorted_keywords = sorted(keyword_colors.keys(), key=len, reverse=True)
            for kw in sorted_keywords:
                light_color = keyword_colors[kw]
                if '\\b' in kw:
                    pattern_str = kw.replace('\\\\b', '\\b')
                    pattern = re.compile(pattern_str, re.IGNORECASE)
                else:
                    pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m, c=light_color: f'<span style="background-color: {c}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        imported_entries = []
        for entry, cats in filtered:
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_str = format_date_nice(date_obj) if date_obj else ""
            highlighted_html = highlight_keywords(full_text, cats or [])

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="progress",
                categories=cats,
                category_colors=CATEGORY_COLORS,
                highlighted_html=highlighted_html,
                filter_callback=self._apply_progress_filter,
                parent_container=getattr(self, 'popup_progress_import_container', None),
            )

            import_layout.addWidget(entry_frame)
            imported_entries.append({"checkbox": cb, "text": full_text, "date": entry.get("date", ""), "categories": cats})

        self.popup_progress_imported_entries = imported_entries

    def _build_popup_managing_risk(self):
        """Build managing risk popup with Current and Historical risk sections."""
        container, layout = self._create_popup_container("managing_risk")

        # Risk types for the form - matches Section 5 risk factors
        self.RISK_TYPES = [
            ("violence_others", "Violence to others"),
            ("violence_property", "Violence to property"),
            ("verbal_aggression", "Verbal aggression"),
            ("substance_misuse", "Substance misuse"),
            ("self_harm", "Self harm"),
            ("self_neglect", "Self neglect"),
            ("stalking", "Stalking"),
            ("threatening_behaviour", "Threatening behaviour"),
            ("sexually_inappropriate", "Sexually inappropriate behaviour"),
            ("vulnerability", "Vulnerability"),
            ("bullying_victimisation", "Bullying/victimisation"),
            ("absconding", "Absconding/AWOL"),
            ("reoffending", "Reoffending"),
        ]

        self._risk_current_widgets = {}
        self._risk_historical_widgets = {}

        slider_style = """
            QSlider::groove:horizontal {
                height: 4px;
                background: #e5e7eb;
                border-radius: 2px;
            }
            QSlider::handle:horizontal {
                width: 14px;
                margin: -5px 0;
                background: #991b1b;
                border-radius: 7px;
            }
            QSlider::sub-page:horizontal {
                background: #991b1b;
                border-radius: 2px;
            }
        """

        # === CURRENT RISK SECTION ===
        current_frame = QFrame()
        current_frame.setStyleSheet("""
            QFrame {
                background: #fef2f2;
                border: 1px solid #fecaca;
                border-radius: 8px;
            }
        """)
        current_layout = QVBoxLayout(current_frame)
        current_layout.setContentsMargins(12, 12, 12, 12)
        current_layout.setSpacing(8)

        current_header = QLabel("Current Risk Factors")
        current_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #991b1b; background: transparent;")
        current_layout.addWidget(current_header)

        for key, label in self.RISK_TYPES:
            row = QWidget()
            row.setStyleSheet("background: transparent;")
            row_layout = QVBoxLayout(row)
            row_layout.setContentsMargins(0, 0, 0, 0)
            row_layout.setSpacing(2)

            cb = QCheckBox(label)
            cb.setStyleSheet("font-size: 17px; background: transparent;")
            row_layout.addWidget(cb)

            # Severity slider (hidden by default)
            slider_container = QWidget()
            slider_container.setStyleSheet("background: transparent;")
            slider_layout = QHBoxLayout(slider_container)
            slider_layout.setContentsMargins(24, 0, 0, 4)
            slider_layout.setSpacing(8)

            severity_lbl = QLabel("Low")
            severity_lbl.setFixedWidth(60)
            severity_lbl.setStyleSheet("font-size: 16px; color: #991b1b; font-weight: 600; background: transparent;")

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setMinimum(1)
            slider.setMaximum(3)
            slider.setValue(1)  # Default to Low
            slider.setFixedWidth(100)
            slider.setStyleSheet(slider_style)
            slider.valueChanged.connect(lambda v, l=severity_lbl: l.setText(["Low", "Medium", "High"][v-1]))

            slider_layout.addWidget(slider)
            slider_layout.addWidget(severity_lbl)
            slider_layout.addStretch()

            slider_container.hide()
            row_layout.addWidget(slider_container)

            cb.toggled.connect(lambda checked, sc=slider_container: sc.setVisible(checked))

            current_layout.addWidget(row)

            self._risk_current_widgets[key] = {
                "checkbox": cb,
                "slider": slider,
                "severity_label": severity_lbl,
                "slider_container": slider_container
            }

        layout.addWidget(current_frame)

        # === HISTORICAL RISK SECTION ===
        historical_frame = QFrame()
        historical_frame.setStyleSheet("""
            QFrame {
                background: #fffbeb;
                border: 1px solid #fcd34d;
                border-radius: 8px;
            }
        """)
        historical_layout = QVBoxLayout(historical_frame)
        historical_layout.setContentsMargins(12, 12, 12, 12)
        historical_layout.setSpacing(8)

        historical_header = QLabel("Historical Risk Factors")
        historical_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #806000; background: transparent;")
        historical_layout.addWidget(historical_header)

        for key, label in self.RISK_TYPES:
            row = QWidget()
            row.setStyleSheet("background: transparent;")
            row_layout = QVBoxLayout(row)
            row_layout.setContentsMargins(0, 0, 0, 0)
            row_layout.setSpacing(2)

            cb = QCheckBox(label)
            cb.setStyleSheet("font-size: 17px; background: transparent;")
            row_layout.addWidget(cb)

            # Severity slider (hidden by default)
            slider_container = QWidget()
            slider_container.setStyleSheet("background: transparent;")
            slider_layout = QHBoxLayout(slider_container)
            slider_layout.setContentsMargins(24, 0, 0, 4)
            slider_layout.setSpacing(8)

            severity_lbl = QLabel("Low")
            severity_lbl.setFixedWidth(60)
            severity_lbl.setStyleSheet("font-size: 16px; color: #806000; font-weight: 600; background: transparent;")

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setMinimum(1)
            slider.setMaximum(3)
            slider.setValue(1)  # Default to Low
            slider.setFixedWidth(100)
            slider.setStyleSheet(slider_style.replace("#991b1b", "#b45309"))
            slider.valueChanged.connect(lambda v, l=severity_lbl: l.setText(["Low", "Medium", "High"][v-1]))

            slider_layout.addWidget(slider)
            slider_layout.addWidget(severity_lbl)
            slider_layout.addStretch()

            slider_container.hide()
            row_layout.addWidget(slider_container)

            cb.toggled.connect(lambda checked, sc=slider_container: sc.setVisible(checked))

            historical_layout.addWidget(row)

            self._risk_historical_widgets[key] = {
                "checkbox": cb,
                "slider": slider,
                "severity_label": severity_lbl,
                "slider_container": slider_container
            }

        layout.addWidget(historical_frame)

        # Sync current risks to historical risks
        def sync_current_to_historical(key):
            """When current risk is changed, sync to historical."""
            current = self._risk_current_widgets[key]
            historical = self._risk_historical_widgets[key]

            # Sync checkbox state
            if current["checkbox"].isChecked():
                historical["checkbox"].setChecked(True)
                # Sync slider value
                historical["slider"].setValue(current["slider"].value())
                historical["severity_label"].setText(current["severity_label"].text())

        # Connect current risk checkboxes and sliders to sync to historical
        for key in self._risk_current_widgets:
            current = self._risk_current_widgets[key]
            # When checkbox is toggled, sync to historical
            current["checkbox"].toggled.connect(lambda checked, k=key: sync_current_to_historical(k) if checked else None)
            # When slider value changes, sync to historical
            current["slider"].valueChanged.connect(lambda v, k=key: sync_current_to_historical(k))

        layout.addStretch()
        self._add_send_button(layout, "managing_risk", self._generate_managing_risk)
        # Connect sliders and checkboxes for preview updates
        risk_widgets = []
        for risk_data in self._risk_current_widgets.values():
            risk_widgets.extend([risk_data["checkbox"], risk_data["slider"]])
        for risk_data in self._risk_historical_widgets.values():
            risk_widgets.extend([risk_data["checkbox"], risk_data["slider"]])
        self._connect_preview_updates("managing_risk", risk_widgets)

    def _sync_risk_to_section9(self, risk_key: str, checked: bool):
        """Sync Section 5 risk factor selection to Section 9 Current and Historical risk factors."""
        # Check if Section 9 widgets exist
        if not hasattr(self, '_risk_current_widgets') or not hasattr(self, '_risk_historical_widgets'):
            return

        # Sync to current risk
        if risk_key in self._risk_current_widgets:
            current_cb = self._risk_current_widgets[risk_key]["checkbox"]
            if current_cb.isChecked() != checked:
                current_cb.setChecked(checked)

        # Sync to historical risk
        if risk_key in self._risk_historical_widgets:
            historical_cb = self._risk_historical_widgets[risk_key]["checkbox"]
            if historical_cb.isChecked() != checked:
                historical_cb.setChecked(checked)

    def _generate_managing_risk(self) -> str:
        """Generate narrative risk text from Current and Historical selections."""
        def build_narrative(risks, is_historical=False):
            if not risks:
                return ""

            # Sort by severity
            sorted_risks = sorted(risks, key=lambda x: x[1], reverse=True)
            severity_words = {1: "low", 2: "moderate", 3: "high"}

            high = [r[0] for r in sorted_risks if r[1] == 3]
            moderate = [r[0] for r in sorted_risks if r[1] == 2]
            low = [r[0] for r in sorted_risks if r[1] == 1]

            def join_list(items):
                if len(items) == 1:
                    return items[0]
                elif len(items) == 2:
                    return f"{items[0]} and {items[1]}"
                else:
                    return ", ".join(items[:-1]) + f", and {items[-1]}"

            parts = []
            prefix = "Historically, the" if is_historical else "The"

            if high:
                parts.append(f"risk of {join_list(high)} is high" if len(high) == 1 else f"risks of {join_list(high)} are high")
            if moderate:
                if parts:
                    parts.append(f"{join_list(moderate)} is moderate" if len(moderate) == 1 else f"{join_list(moderate)} are moderate")
                else:
                    parts.append(f"risk of {join_list(moderate)} is moderate" if len(moderate) == 1 else f"risks of {join_list(moderate)} are moderate")
            if low:
                if parts:
                    parts.append(f"{join_list(low)} is low" if len(low) == 1 else f"{join_list(low)} are low")
                else:
                    parts.append(f"risk of {join_list(low)} is low" if len(low) == 1 else f"risks of {join_list(low)} are low")

            if len(parts) == 1:
                return f"{prefix} {parts[0]}."
            elif len(parts) == 2:
                return f"{prefix} {parts[0]}, and {parts[1]}."
            else:
                return f"{prefix} {parts[0]}, {parts[1]}, and {parts[2]}."

        # Gather current risks
        current_risks = []
        for key, widgets in self._risk_current_widgets.items():
            if widgets["checkbox"].isChecked():
                risk_name = widgets["checkbox"].text().lower()
                severity_val = widgets["slider"].value()
                current_risks.append((risk_name, severity_val))

        # Gather historical risks
        historical_risks = []
        for key, widgets in self._risk_historical_widgets.items():
            if widgets["checkbox"].isChecked():
                risk_name = widgets["checkbox"].text().lower()
                severity_val = widgets["slider"].value()
                historical_risks.append((risk_name, severity_val))

        sections = []
        if current_risks:
            sections.append(build_narrative(current_risks, is_historical=False))
        if historical_risks:
            sections.append(build_narrative(historical_risks, is_historical=True))

        return "\n\n".join(sections) if sections else "No specific risk factors identified."

    def _build_popup_risk_addressed(self):
        """Build risk addressed popup with structured prompts from template and imported notes."""
        container, layout = self._create_popup_container_with_imports("risk_addressed")

        label_style = "font-size: 17px;"
        radio_style = "QRadioButton { font-size: 17px; }"
        text_style = """
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 8px;
                font-size: 17px;
            }
            QTextEdit:focus {
                border-color: #991b1b;
            }
        """

        # 1. Progress and concerns
        lbl1 = QLabel("<b>1. Progress and Issues of Concern:</b>")
        lbl1.setStyleSheet(label_style)
        lbl1.setWordWrap(True)
        layout.addWidget(lbl1)
        self.popup_risk_progress = QTextEdit()
        self.popup_risk_progress.setMaximumHeight(70)
        self.popup_risk_progress.setPlaceholderText("Describe the progress the patient has made (hospital and/or community) and any issues of concern...")
        self.popup_risk_progress.setStyleSheet(text_style)
        layout.addWidget(self.popup_risk_progress)
        layout.addSpacing(6)

        # 2. Understanding factors
        lbl2 = QLabel("<b>2. Factors Underpinning Index Offence:</b>")
        lbl2.setStyleSheet(label_style)
        lbl2.setWordWrap(True)
        layout.addWidget(lbl2)
        self.popup_risk_factors = QTextEdit()
        self.popup_risk_factors.setMaximumHeight(70)
        self.popup_risk_factors.setPlaceholderText("What is the team's current understanding of the factors underpinning the index offence and previous dangerous behaviour...")
        self.popup_risk_factors.setStyleSheet(text_style)
        layout.addWidget(self.popup_risk_factors)
        layout.addSpacing(6)

        # 3. Attitudes to index offence
        lbl3 = QLabel("<b>3. Attitudes to Index Offence & Victims:</b>")
        lbl3.setStyleSheet(label_style)
        lbl3.setWordWrap(True)
        layout.addWidget(lbl3)
        self.popup_risk_attitudes = QTextEdit()
        self.popup_risk_attitudes.setMaximumHeight(70)
        self.popup_risk_attitudes.setPlaceholderText("What are the patient's current attitudes to the index offence, other dangerous behaviour and any previous victims...")
        self.popup_risk_attitudes.setStyleSheet(text_style)
        layout.addWidget(self.popup_risk_attitudes)
        layout.addSpacing(6)

        # 4. Prevent referral
        lbl4 = QLabel("<b>4. Prevent Referral:</b>")
        lbl4.setStyleSheet(label_style)
        layout.addWidget(lbl4)
        prevent_row = QHBoxLayout()
        self.popup_prevent_group = QButtonGroup(self)
        self.popup_prevent_yes = QRadioButton("Yes")
        self.popup_prevent_yes.setStyleSheet(radio_style)
        self.popup_prevent_no = QRadioButton("No")
        self.popup_prevent_no.setStyleSheet(radio_style)
        self.popup_prevent_na = QRadioButton("N/A")
        self.popup_prevent_na.setStyleSheet(radio_style)
        self.popup_prevent_group.addButton(self.popup_prevent_yes)
        self.popup_prevent_group.addButton(self.popup_prevent_no)
        self.popup_prevent_group.addButton(self.popup_prevent_na)
        referred_lbl = QLabel("Referred to Prevent?")
        referred_lbl.setStyleSheet(label_style)
        prevent_row.addWidget(referred_lbl)
        prevent_row.addWidget(self.popup_prevent_yes)
        prevent_row.addWidget(self.popup_prevent_no)
        prevent_row.addWidget(self.popup_prevent_na)
        prevent_row.addStretch()
        layout.addLayout(prevent_row)

        self.popup_prevent_outcome = QLineEdit()
        self.popup_prevent_outcome.setPlaceholderText("If referred, outcome of referral...")
        self.popup_prevent_outcome.setStyleSheet("padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;")
        layout.addWidget(self.popup_prevent_outcome)

        layout.addStretch()
        self._add_send_button(layout, "risk_addressed", self._generate_risk_addressed)
        self._connect_preview_updates("risk_addressed", [
            self.popup_risk_progress, self.popup_risk_factors, self.popup_risk_attitudes,
            self.popup_prevent_yes, self.popup_prevent_no, self.popup_prevent_na, self.popup_prevent_outcome
        ])

    def _generate_risk_addressed(self) -> str:
        """Generate risk addressed text from template fields."""
        parts = []

        progress = self.popup_risk_progress.toPlainText().strip()
        if progress:
            parts.append(f"Progress and issues of concern: {progress}")

        factors = self.popup_risk_factors.toPlainText().strip()
        if factors:
            parts.append(f"Factors underpinning index offence: {factors}")

        attitudes = self.popup_risk_attitudes.toPlainText().strip()
        if attitudes:
            parts.append(f"Attitudes to index offence and victims: {attitudes}")

        # Prevent referral
        if self.popup_prevent_yes.isChecked():
            outcome = self.popup_prevent_outcome.text().strip()
            if outcome:
                parts.append(f"The patient has been referred to Prevent. Outcome: {outcome}")
            else:
                parts.append("The patient has been referred to Prevent.")
        elif self.popup_prevent_no.isChecked():
            parts.append("The patient has not been referred to Prevent.")
        elif hasattr(self, 'popup_prevent_na') and self.popup_prevent_na.isChecked():
            parts.append("Prevent is not applicable in this case.")

        result = "\n".join(parts) if parts else ""

        # Imported notes (checked entries)
        imported_texts = []
        if hasattr(self, 'popup_risk_addressed_imported_entries'):
            for entry in self.popup_risk_addressed_imported_entries:
                if entry["checkbox"].isChecked():
                    date = entry.get("date", "")
                    text = entry["text"]
                    if date:
                        imported_texts.append(f"[{date}] {text}")
                    else:
                        imported_texts.append(text)

        if imported_texts:
            if result:
                result += "\n\n--- Imported Notes ---\n" + "\n".join(imported_texts)
            else:
                result = "--- Imported Notes ---\n" + "\n".join(imported_texts)

        return result if result else "[No information provided]"

    def _build_popup_abscond(self):
        """Build abscond popup with imported notes section."""
        container, layout = self._create_popup_container_with_imports("abscond")

        label_style = "font-size: 17px; color: #374151;"
        radio_style = "QRadioButton { font-size: 17px; }"

        # Get the import layout that was stored by _create_popup_container_with_imports
        import_layout = getattr(self, 'popup_abscond_import_layout', None)

        # Filter dropdown for AWOL categories (add to import section if available)
        if import_layout:
            filter_row = QHBoxLayout()
            filter_label = QLabel("Filter:")
            filter_label.setStyleSheet("font-size: 16px; color: #6b7280; background: transparent;")
            self.popup_abscond_filter = QComboBox()
            self.popup_abscond_filter.addItem("All")
            self.popup_abscond_filter.setStyleSheet("font-size: 16px; padding: 2px;")
            self.popup_abscond_filter.currentTextChanged.connect(self._apply_abscond_filter)
            filter_row.addWidget(filter_label)
            filter_row.addWidget(self.popup_abscond_filter)
            filter_row.addStretch()
            import_layout.insertLayout(0, filter_row)

        # AWOL Yes/No
        awol_row = QHBoxLayout()
        awol_lbl = QLabel("Any AWOL incidents?")
        awol_lbl.setStyleSheet(label_style)
        awol_row.addWidget(awol_lbl)
        self.popup_awol_group = QButtonGroup(self)
        self.popup_awol_yes = QRadioButton("Yes")
        self.popup_awol_yes.setStyleSheet(radio_style)
        self.popup_awol_no = QRadioButton("No")
        self.popup_awol_no.setStyleSheet(radio_style)
        self.popup_awol_no.setChecked(True)
        self.popup_awol_group.addButton(self.popup_awol_yes)
        self.popup_awol_group.addButton(self.popup_awol_no)
        awol_row.addWidget(self.popup_awol_yes)
        awol_row.addWidget(self.popup_awol_no)
        awol_row.addStretch()
        layout.addLayout(awol_row)

        details_lbl = QLabel("Details:")
        details_lbl.setStyleSheet(label_style)
        layout.addWidget(details_lbl)
        self.popup_abscond_details = QTextEdit()
        self.popup_abscond_details.setMaximumHeight(120)
        self.popup_abscond_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;")
        layout.addWidget(self.popup_abscond_details)

        layout.addStretch()
        self._add_send_button(layout, "abscond", self._generate_abscond)
        self._connect_preview_updates("abscond", [
            self.popup_awol_yes, self.popup_awol_no, self.popup_abscond_details
        ])

    def _generate_abscond(self) -> str:
        if self.popup_awol_no.isChecked():
            return "There have been no AWOL concerns in the last 12 months."
        details = self.popup_abscond_details.toPlainText()
        return f"There have been AWOL concerns. {details}" if details else "There have been AWOL concerns."

    def _build_popup_mappa(self):
        """Build MAPPA popup matching template structure with imported notes."""
        container, layout = self._create_popup_container_with_imports("mappa")

        # Get the import layout for the filter dropdown
        import_layout = getattr(self, 'popup_mappa_import_layout', None)

        label_style = "font-size: 17px;"
        radio_style = "QRadioButton { font-size: 17px; }"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        text_style = """
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 8px;
                font-size: 17px;
            }
        """

        # 1. MAPPA Category
        lbl1 = QLabel("<b>1. MAPPA Category:</b>")
        lbl1.setStyleSheet(label_style)
        layout.addWidget(lbl1)
        self.popup_mappa_cat = QComboBox()
        self.popup_mappa_cat.addItems(["Not applicable", "Category 1", "Category 2", "Category 3"])
        self.popup_mappa_cat.setStyleSheet(input_style)
        layout.addWidget(self.popup_mappa_cat)
        layout.addSpacing(8)

        # 2. MAPPA Level
        lbl2 = QLabel("<b>2. Level at which managed:</b>")
        lbl2.setStyleSheet(label_style)
        layout.addWidget(lbl2)
        self.popup_mappa_level = QComboBox()
        self.popup_mappa_level.addItems(["N/A", "Level 1", "Level 2", "Level 3"])
        self.popup_mappa_level.setStyleSheet(input_style)
        layout.addWidget(self.popup_mappa_level)
        layout.addSpacing(8)

        # 3. Date of referral
        lbl3 = QLabel("<b>3. Date of referral to MAPPA:</b>")
        lbl3.setStyleSheet(label_style)
        layout.addWidget(lbl3)

        # Date known/not known radio buttons
        mappa_date_row = QHBoxLayout()
        mappa_date_row.setSpacing(12)

        self.popup_mappa_date_group = QButtonGroup(self)
        self.popup_mappa_date_known = QRadioButton("Date known")
        self.popup_mappa_date_known.setStyleSheet(radio_style)
        self.popup_mappa_date_unknown = QRadioButton("Not known")
        self.popup_mappa_date_unknown.setStyleSheet(radio_style)
        self.popup_mappa_date_group.addButton(self.popup_mappa_date_known)
        self.popup_mappa_date_group.addButton(self.popup_mappa_date_unknown)
        self.popup_mappa_date_known.setChecked(True)

        mappa_date_row.addWidget(self.popup_mappa_date_known)
        mappa_date_row.addWidget(self.popup_mappa_date_unknown)
        mappa_date_row.addStretch()
        layout.addLayout(mappa_date_row)

        self.popup_mappa_date = QDateEdit()
        self.popup_mappa_date.setCalendarPopup(True)
        self.popup_mappa_date.setDisplayFormat("dd/MM/yyyy")
        self.popup_mappa_date.setDate(QDate.currentDate())
        self.popup_mappa_date.setStyleSheet(input_style)
        layout.addWidget(self.popup_mappa_date)

        # Connect radio to enable/disable date field
        self.popup_mappa_date_unknown.toggled.connect(
            lambda checked: self.popup_mappa_date.setEnabled(not checked)
        )
        self.popup_mappa_date_unknown.toggled.connect(
            lambda checked: self.popup_mappa_date.setStyleSheet(
                input_style + " background: #f3f4f6; color: #9ca3af;" if checked else input_style
            )
        )
        self.popup_mappa_date_known.toggled.connect(lambda: self._update_preview("mappa"))
        self.popup_mappa_date_unknown.toggled.connect(lambda: self._update_preview("mappa"))

        layout.addSpacing(8)

        # 4. MAPPA comments
        lbl4 = QLabel("<b>4. MAPPA comments (last 12 months):</b>")
        lbl4.setStyleSheet(label_style)
        layout.addWidget(lbl4)
        self.popup_mappa_comments = QTextEdit()
        self.popup_mappa_comments.setMaximumHeight(80)
        self.popup_mappa_comments.setPlaceholderText("Provide MAPPA comments received in the last 12 months...")
        self.popup_mappa_comments.setStyleSheet(text_style)
        layout.addWidget(self.popup_mappa_comments)
        layout.addSpacing(8)

        # 5. MAPPA Coordinator
        lbl5 = QLabel("<b>5. MAPPA Coordinator name and contact:</b>")
        lbl5.setStyleSheet(label_style)
        layout.addWidget(lbl5)
        self.popup_mappa_coord = QLineEdit()
        self.popup_mappa_coord.setPlaceholderText("Name and contact details of MAPPA coordinator...")
        self.popup_mappa_coord.setStyleSheet(input_style)
        layout.addWidget(self.popup_mappa_coord)

        layout.addStretch()
        self._add_send_button(layout, "mappa", self._generate_mappa)
        self._connect_preview_updates("mappa", [
            self.popup_mappa_cat, self.popup_mappa_level, self.popup_mappa_date,
            self.popup_mappa_comments, self.popup_mappa_coord
        ])

    def _generate_mappa(self) -> str:
        cat = self.popup_mappa_cat.currentText()
        level = self.popup_mappa_level.currentText()
        date_unknown = hasattr(self, 'popup_mappa_date_unknown') and self.popup_mappa_date_unknown.isChecked()
        ref_date = "Not known" if date_unknown else self.popup_mappa_date.date().toString("dd/MM/yyyy")
        comments = self.popup_mappa_comments.toPlainText().strip()
        coord = self.popup_mappa_coord.text().strip()

        if cat == "Not applicable":
            result = "Patient is not currently referred to MAPPA."
        else:
            parts = [f"The patient is managed under MAPPA {cat}, {level}."]
            parts.append(f"Date of referral: {ref_date}.")
            if comments:
                parts.append(f"MAPPA comments: {comments}")
            if coord:
                parts.append(f"MAPPA Coordinator: {coord}")
            result = "\n".join(parts)

        # Imported notes (checked entries)
        imported_texts = []
        if hasattr(self, 'popup_mappa_imported_entries'):
            for entry in self.popup_mappa_imported_entries:
                if entry["checkbox"].isChecked():
                    date = entry.get("date", "")
                    text = entry["text"]
                    if date:
                        imported_texts.append(f"[{date}] {text}")
                    else:
                        imported_texts.append(text)

        if imported_texts:
            if result:
                result += "\n\n--- Imported Notes ---\n" + "\n".join(imported_texts)
            else:
                result = "--- Imported Notes ---\n" + "\n".join(imported_texts)

        return result if result else "[No MAPPA information provided]"

    def _build_popup_victims(self):
        """Build victims popup matching template structure."""
        container, layout = self._create_popup_container("victims")

        label_style = "font-size: 17px;"
        radio_style = "QRadioButton { font-size: 17px; }"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        text_style = """
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 8px;
                font-size: 17px;
            }
        """

        # 1. VLO name and contact
        lbl1 = QLabel("<b>1. Victim Liaison Officer (VLO) name and contact:</b>")
        lbl1.setStyleSheet(label_style)
        layout.addWidget(lbl1)
        self.popup_vlo = QLineEdit()
        self.popup_vlo.setPlaceholderText("Please provide full contact details...")
        self.popup_vlo.setStyleSheet(input_style)
        layout.addWidget(self.popup_vlo)
        layout.addSpacing(8)

        # 2. Date of last contact
        lbl2 = QLabel("<b>2. Date of last discussion/contact with VLO:</b>")
        lbl2.setStyleSheet(label_style)
        layout.addWidget(lbl2)

        # Date known/not known radio buttons
        vlo_date_row = QHBoxLayout()
        vlo_date_row.setSpacing(12)

        self.popup_vlo_date_group = QButtonGroup(self)
        self.popup_vlo_date_known = QRadioButton("Date known")
        self.popup_vlo_date_known.setStyleSheet(radio_style)
        self.popup_vlo_date_unknown = QRadioButton("Not known")
        self.popup_vlo_date_unknown.setStyleSheet(radio_style)
        self.popup_vlo_date_group.addButton(self.popup_vlo_date_known)
        self.popup_vlo_date_group.addButton(self.popup_vlo_date_unknown)
        self.popup_vlo_date_known.setChecked(True)

        vlo_date_row.addWidget(self.popup_vlo_date_known)
        vlo_date_row.addWidget(self.popup_vlo_date_unknown)
        vlo_date_row.addStretch()
        layout.addLayout(vlo_date_row)

        self.popup_vlo_date = QDateEdit()
        self.popup_vlo_date.setCalendarPopup(True)
        self.popup_vlo_date.setDisplayFormat("dd/MM/yyyy")
        self.popup_vlo_date.setDate(QDate.currentDate())
        self.popup_vlo_date.setStyleSheet(input_style)
        layout.addWidget(self.popup_vlo_date)

        # Connect radio to enable/disable date field
        self.popup_vlo_date_unknown.toggled.connect(
            lambda checked: self.popup_vlo_date.setEnabled(not checked)
        )
        self.popup_vlo_date_unknown.toggled.connect(
            lambda checked: self.popup_vlo_date.setStyleSheet(
                input_style + " background: #f3f4f6; color: #9ca3af;" if checked else input_style
            )
        )
        self.popup_vlo_date_known.toggled.connect(lambda: self._update_preview("victims"))
        self.popup_vlo_date_unknown.toggled.connect(lambda: self._update_preview("victims"))

        layout.addSpacing(8)

        # 3. Victim-related concerns
        lbl3 = QLabel("<b>3. Victim-related concerns (last 12 months):</b>")
        lbl3.setStyleSheet(label_style)
        layout.addWidget(lbl3)
        self.popup_victim_concerns = QTextEdit()
        self.popup_victim_concerns.setMaximumHeight(100)
        self.popup_victim_concerns.setPlaceholderText("Have there been any victim-related concerns or contact with the VLO in the last 12 months...")
        self.popup_victim_concerns.setStyleSheet(text_style)
        layout.addWidget(self.popup_victim_concerns)

        layout.addStretch()
        self._add_send_button(layout, "victims", self._generate_victims)
        self._connect_preview_updates("victims", [
            self.popup_vlo, self.popup_vlo_date, self.popup_victim_concerns
        ])

    def _generate_victims(self) -> str:
        vlo = self.popup_vlo.text().strip()
        date_unknown = hasattr(self, 'popup_vlo_date_unknown') and self.popup_vlo_date_unknown.isChecked()
        contact_date = "Not known" if date_unknown else self.popup_vlo_date.date().toString("dd/MM/yyyy")
        concerns = self.popup_victim_concerns.toPlainText().strip()

        parts = []
        if vlo:
            parts.append(f"Victim Liaison Officer: {vlo}.")
        parts.append(f"Date of last contact with VLO: {contact_date}.")
        if concerns:
            parts.append(f"Victim-related concerns (last 12 months): {concerns}")

        return "\n".join(parts) if parts else "No victim liaison information available."

    def _build_popup_leave_report(self):
        """Build leave report popup - 3g style from MOJ Leave form with escorted/unescorted toggle."""
        container, layout = self._create_popup_container_with_imports("leave_report")

        # Initialize state storage for escorted/unescorted entries
        self._leave_escorted_state = {}
        self._leave_unescorted_state = {}
        self._leave_switching = False  # Flag to prevent recursive updates

        slider_style = """
            QSlider::groove:horizontal { background: #c7d2fe; height: 6px; border-radius: 3px; }
            QSlider::handle:horizontal { background: #4f46e5; width: 14px; margin: -4px 0; border-radius: 7px; }
            QSlider:disabled::groove:horizontal { background: #e5e7eb; }
            QSlider:disabled::handle:horizontal { background: #9ca3af; }
        """

        dropdown_style = """
            QComboBox {
                background: white;
                padding: 6px 10px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                font-size: 16px;
            }
            QComboBox::drop-down {
                width: 20px;
            }
        """

        # Escorted/Unescorted radio at the top
        escort_frame = QFrame()
        escort_frame.setStyleSheet("QFrame { background: #dbeafe; border: 2px solid #3b82f6; border-radius: 8px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        escort_layout = QHBoxLayout(escort_frame)
        escort_layout.setContentsMargins(12, 8, 12, 8)
        escort_layout.setSpacing(16)

        escort_label = QLabel("<b>Leave Type:</b>")
        escort_label.setStyleSheet("font-size: 17px; color: #1e40af;")
        escort_layout.addWidget(escort_label)

        self.popup_leave_escort_group = QButtonGroup(self)
        self.popup_leave_escorted = QRadioButton("Escorted")
        self.popup_leave_unescorted = QRadioButton("Unescorted")
        self.popup_leave_escorted.setStyleSheet("font-weight: 600; font-size: 17px; color: #1e40af;")
        self.popup_leave_unescorted.setStyleSheet("font-weight: 600; font-size: 17px; color: #1e40af;")
        self.popup_leave_escort_group.addButton(self.popup_leave_escorted)
        self.popup_leave_escort_group.addButton(self.popup_leave_unescorted)
        self.popup_leave_escorted.setChecked(True)  # Default to escorted
        escort_layout.addWidget(self.popup_leave_escorted)
        escort_layout.addWidget(self.popup_leave_unescorted)
        escort_layout.addStretch()

        layout.addWidget(escort_frame)

        # Row 1: Leaves, Frequency, Duration
        row1 = QHBoxLayout()
        row1.setSpacing(8)

        row1.addWidget(QLabel("Leaves:"))
        self.popup_leaves_dropdown = QComboBox()
        self.popup_leaves_dropdown.addItems([str(i) for i in range(1, 8)])
        self.popup_leaves_dropdown.setMinimumWidth(60)
        self.popup_leaves_dropdown.setStyleSheet(dropdown_style)
        row1.addWidget(self.popup_leaves_dropdown)

        row1.addWidget(QLabel("Frequency:"))
        self.popup_frequency_dropdown = QComboBox()
        self.popup_frequency_dropdown.addItems(["Weekly", "2 weekly", "3 weekly", "Monthly", "2 monthly"])
        self.popup_frequency_dropdown.setMinimumWidth(100)
        self.popup_frequency_dropdown.setStyleSheet(dropdown_style)
        row1.addWidget(self.popup_frequency_dropdown)

        row1.addWidget(QLabel("Duration:"))
        self.popup_duration_dropdown = QComboBox()
        self.popup_duration_dropdown.addItems(["30 mins", "1 hour", "2 hours", "3 hours", "4 hours", "5 hours", "6 hours", "7 hours", "8 hours"])
        self.popup_duration_dropdown.setMinimumWidth(80)
        self.popup_duration_dropdown.setStyleSheet(dropdown_style)
        row1.addWidget(self.popup_duration_dropdown)

        row1.addStretch()
        layout.addLayout(row1)

        # Leave types with individual linked sliders
        weight_frame = QFrame()
        weight_frame.setStyleSheet("QFrame { background: #e0e7ff; border: 1px solid #6366f1; border-radius: 6px; } QLabel { background: transparent; border: none; } QCheckBox { background: transparent; border: none; }")
        weight_layout = QVBoxLayout(weight_frame)
        weight_layout.setContentsMargins(10, 8, 10, 8)
        weight_layout.setSpacing(6)

        weight_layout.addWidget(QLabel("<b>Leave types & weighting:</b>"))

        self.popup_leave_type_widgets = {}
        self._popup_updating_sliders = False

        leave_types = [
            ("ground", "Ground"),
            ("local", "Local community"),
            ("community", "Community"),
            ("extended", "Extended community"),
            ("overnight", "Overnight")
        ]

        for key, label in leave_types:
            row = QHBoxLayout()
            row.setSpacing(8)

            cb = QCheckBox(label)
            cb.setFixedWidth(130)
            row.addWidget(cb)

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setMinimum(0)
            slider.setMaximum(100)
            slider.setValue(0)
            slider.setEnabled(False)
            slider.setStyleSheet(slider_style)
            row.addWidget(slider, 1)

            pct_label = QLabel("0%")
            pct_label.setFixedWidth(35)
            pct_label.setStyleSheet("font-weight: 600; color: #4f46e5;")
            row.addWidget(pct_label)

            self.popup_leave_type_widgets[key] = {"cb": cb, "slider": slider, "label": pct_label}

            cb.stateChanged.connect(lambda state, k=key: self._on_popup_leave_type_toggled(k, state))
            slider.valueChanged.connect(lambda val, k=key: self._on_popup_leave_slider_changed(k, val))

            weight_layout.addLayout(row)

        layout.addWidget(weight_frame)

        # Other leave types
        other_frame = QFrame()
        other_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #f59e0b; border-radius: 6px; } QLabel { background: transparent; border: none; } QCheckBox { background: transparent; border: none; }")
        other_layout = QVBoxLayout(other_frame)
        other_layout.setContentsMargins(10, 8, 10, 8)
        other_layout.setSpacing(6)

        other_layout.addWidget(QLabel("<b>Other leave:</b>"))
        other_row = QHBoxLayout()
        self.popup_leave_medical_cb = QCheckBox("Medical")
        self.popup_leave_court_cb = QCheckBox("Court")
        self.popup_leave_compassionate_cb = QCheckBox("Compassionate")
        for cb in [self.popup_leave_medical_cb, self.popup_leave_court_cb, self.popup_leave_compassionate_cb]:
            cb.stateChanged.connect(lambda: self._update_preview("leave_report"))
            other_row.addWidget(cb)
        other_row.addStretch()
        other_layout.addLayout(other_row)

        layout.addWidget(other_frame)

        # Leave suspended
        suspended_frame = QFrame()
        suspended_frame.setStyleSheet("QFrame { background: #fee2e2; border: 1px solid #ef4444; border-radius: 6px; } QLabel { background: transparent; border: none; } QRadioButton { background: transparent; border: none; }")
        suspended_layout = QVBoxLayout(suspended_frame)
        suspended_layout.setContentsMargins(10, 8, 10, 8)
        suspended_layout.setSpacing(6)

        suspended_layout.addWidget(QLabel("<b>Leave ever suspended:</b>"))
        self.popup_suspended_group = QButtonGroup(self)
        suspended_row = QHBoxLayout()
        self.popup_suspended_yes = QRadioButton("Yes")
        self.popup_suspended_no = QRadioButton("No")
        self.popup_suspended_group.addButton(self.popup_suspended_yes)
        self.popup_suspended_group.addButton(self.popup_suspended_no)
        suspended_row.addWidget(self.popup_suspended_yes)
        suspended_row.addWidget(self.popup_suspended_no)
        suspended_row.addStretch()
        suspended_layout.addLayout(suspended_row)

        # Suspension details (shown when Yes)
        self.popup_suspension_details_container = QFrame()
        self.popup_suspension_details_container.setStyleSheet("QFrame { background: transparent; } QLabel { background: transparent; border: none; }")
        susp_details_layout = QVBoxLayout(self.popup_suspension_details_container)
        susp_details_layout.setContentsMargins(0, 4, 0, 0)
        susp_details_layout.setSpacing(4)
        susp_details_layout.addWidget(QLabel("Details:"))
        self.popup_suspension_details = QTextEdit()
        self.popup_suspension_details.setPlaceholderText("Describe the circumstances of leave suspension...")
        self.popup_suspension_details.setMinimumHeight(60)
        self.popup_suspension_details.setMaximumHeight(80)
        self.popup_suspension_details.setStyleSheet("QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 4px; padding: 6px; font-size: 15px; }")
        self.popup_suspension_details.textChanged.connect(lambda: self._update_preview("leave_report"))
        susp_details_layout.addWidget(self.popup_suspension_details)
        self.popup_suspension_details_container.setVisible(False)
        suspended_layout.addWidget(self.popup_suspension_details_container)

        def on_suspended_changed():
            self.popup_suspension_details_container.setVisible(self.popup_suspended_yes.isChecked())
            self._update_preview("leave_report")

        self.popup_suspended_yes.toggled.connect(on_suspended_changed)
        self.popup_suspended_no.toggled.connect(on_suspended_changed)

        layout.addWidget(suspended_frame)

        # Connect escorted/unescorted toggle to save/restore state
        def on_escort_type_changed():
            if self._leave_switching:
                return
            self._leave_switching = True
            try:
                if self.popup_leave_escorted.isChecked():
                    # Save unescorted state, restore escorted state
                    self._save_leave_state(self._leave_unescorted_state)
                    self._restore_leave_state(self._leave_escorted_state)
                else:
                    # Save escorted state, restore unescorted state
                    self._save_leave_state(self._leave_escorted_state)
                    self._restore_leave_state(self._leave_unescorted_state)
                self._update_preview("leave_report")
            finally:
                self._leave_switching = False

        self.popup_leave_escorted.toggled.connect(on_escort_type_changed)

        layout.addStretch()
        self._add_send_button(layout, "leave_report", self._generate_leave_report)
        self._connect_preview_updates("leave_report", [
            self.popup_leaves_dropdown, self.popup_frequency_dropdown, self.popup_duration_dropdown,
            self.popup_leave_medical_cb, self.popup_leave_court_cb, self.popup_leave_compassionate_cb,
            self.popup_suspended_yes, self.popup_suspended_no
        ])

    def _save_leave_state(self, state_dict):
        """Save current leave report form state to a dictionary."""
        state_dict['leaves'] = self.popup_leaves_dropdown.currentIndex()
        state_dict['frequency'] = self.popup_frequency_dropdown.currentIndex()
        state_dict['duration'] = self.popup_duration_dropdown.currentIndex()
        state_dict['leave_types'] = {}
        for key, widgets in self.popup_leave_type_widgets.items():
            state_dict['leave_types'][key] = {
                'checked': widgets['cb'].isChecked(),
                'value': widgets['slider'].value()
            }
        state_dict['medical'] = self.popup_leave_medical_cb.isChecked()
        state_dict['court'] = self.popup_leave_court_cb.isChecked()
        state_dict['compassionate'] = self.popup_leave_compassionate_cb.isChecked()
        state_dict['suspended_yes'] = self.popup_suspended_yes.isChecked()
        state_dict['suspended_no'] = self.popup_suspended_no.isChecked()
        state_dict['suspension_details'] = self.popup_suspension_details.toPlainText()

    def _restore_leave_state(self, state_dict):
        """Restore leave report form state from a dictionary."""
        if not state_dict:
            # Reset to defaults
            self.popup_leaves_dropdown.setCurrentIndex(0)
            self.popup_frequency_dropdown.setCurrentIndex(0)
            self.popup_duration_dropdown.setCurrentIndex(0)
            for widgets in self.popup_leave_type_widgets.values():
                widgets['cb'].setChecked(False)
                widgets['slider'].setValue(0)
            self.popup_leave_medical_cb.setChecked(False)
            self.popup_leave_court_cb.setChecked(False)
            self.popup_leave_compassionate_cb.setChecked(False)
            self.popup_suspended_yes.setChecked(False)
            self.popup_suspended_no.setChecked(False)
            self.popup_suspension_details.clear()
            return

        self.popup_leaves_dropdown.setCurrentIndex(state_dict.get('leaves', 0))
        self.popup_frequency_dropdown.setCurrentIndex(state_dict.get('frequency', 0))
        self.popup_duration_dropdown.setCurrentIndex(state_dict.get('duration', 0))
        leave_types = state_dict.get('leave_types', {})
        for key, widgets in self.popup_leave_type_widgets.items():
            lt = leave_types.get(key, {})
            widgets['cb'].setChecked(lt.get('checked', False))
            widgets['slider'].setValue(lt.get('value', 0))
        self.popup_leave_medical_cb.setChecked(state_dict.get('medical', False))
        self.popup_leave_court_cb.setChecked(state_dict.get('court', False))
        self.popup_leave_compassionate_cb.setChecked(state_dict.get('compassionate', False))
        if state_dict.get('suspended_yes'):
            self.popup_suspended_yes.setChecked(True)
        elif state_dict.get('suspended_no'):
            self.popup_suspended_no.setChecked(True)
        else:
            self.popup_suspended_yes.setChecked(False)
            self.popup_suspended_no.setChecked(False)
        self.popup_suspension_details.setPlainText(state_dict.get('suspension_details', ''))

    def _on_popup_leave_type_toggled(self, key, state):
        """Handle leave type checkbox toggle - enable/disable slider and redistribute weights."""
        if not hasattr(self, 'popup_leave_type_widgets'):
            return

        widgets = self.popup_leave_type_widgets[key]
        is_checked = state == Qt.CheckState.Checked.value

        widgets["slider"].setEnabled(is_checked)

        if is_checked:
            self._add_popup_leave_type_weight(key)
        else:
            self._remove_popup_leave_type_weight(key)

        self._update_preview("leave_report")

    def _add_popup_leave_type_weight(self, new_key):
        """Add a new leave type, taking weight proportionally from existing checked items."""
        self._popup_updating_sliders = True

        other_checked = [(k, w["slider"].value()) for k, w in self.popup_leave_type_widgets.items()
                         if w["cb"].isChecked() and k != new_key]

        if not other_checked:
            slider = self.popup_leave_type_widgets[new_key]["slider"]
            slider.blockSignals(True)
            slider.setValue(100)
            slider.blockSignals(False)
            self.popup_leave_type_widgets[new_key]["label"].setText("100%")
        else:
            new_item_share = 1
            remaining = 100 - new_item_share

            total_existing = sum(v for _, v in other_checked)
            if total_existing == 0:
                total_existing = 100

            for k, old_val in other_checked:
                if total_existing > 0:
                    new_val = int((old_val / total_existing) * remaining)
                else:
                    new_val = remaining // len(other_checked)
                slider = self.popup_leave_type_widgets[k]["slider"]
                slider.blockSignals(True)
                slider.setValue(new_val)
                slider.blockSignals(False)
                self.popup_leave_type_widgets[k]["label"].setText(f"{new_val}%")

            slider = self.popup_leave_type_widgets[new_key]["slider"]
            slider.blockSignals(True)
            slider.setValue(new_item_share)
            slider.blockSignals(False)
            self.popup_leave_type_widgets[new_key]["label"].setText(f"{new_item_share}%")

        self._popup_updating_sliders = False

    def _remove_popup_leave_type_weight(self, removed_key):
        """Remove a leave type, distributing its weight to remaining checked items."""
        self._popup_updating_sliders = True

        removed_weight = self.popup_leave_type_widgets[removed_key]["slider"].value()

        self.popup_leave_type_widgets[removed_key]["slider"].blockSignals(True)
        self.popup_leave_type_widgets[removed_key]["slider"].setValue(0)
        self.popup_leave_type_widgets[removed_key]["slider"].blockSignals(False)
        self.popup_leave_type_widgets[removed_key]["label"].setText("0%")

        other_checked = [(k, w["slider"].value()) for k, w in self.popup_leave_type_widgets.items()
                         if w["cb"].isChecked() and k != removed_key]

        if other_checked:
            total = sum(v for _, v in other_checked)
            if total > 0:
                for k, old_val in other_checked:
                    new_val = int((old_val / total) * 100)
                    self.popup_leave_type_widgets[k]["slider"].blockSignals(True)
                    self.popup_leave_type_widgets[k]["slider"].setValue(new_val)
                    self.popup_leave_type_widgets[k]["slider"].blockSignals(False)
                    self.popup_leave_type_widgets[k]["label"].setText(f"{new_val}%")

        self._popup_updating_sliders = False

    def _on_popup_leave_slider_changed(self, key, value):
        """Handle slider value change - redistribute remaining weight."""
        if self._popup_updating_sliders:
            return

        self._popup_updating_sliders = True
        self.popup_leave_type_widgets[key]["label"].setText(f"{value}%")

        other_checked = [(k, w["slider"].value()) for k, w in self.popup_leave_type_widgets.items()
                         if w["cb"].isChecked() and k != key]

        if other_checked:
            remaining = 100 - value
            total = sum(v for _, v in other_checked)
            if total > 0:
                for k, old_val in other_checked:
                    new_val = int((old_val / total) * remaining)
                    self.popup_leave_type_widgets[k]["slider"].blockSignals(True)
                    self.popup_leave_type_widgets[k]["slider"].setValue(new_val)
                    self.popup_leave_type_widgets[k]["slider"].blockSignals(False)
                    self.popup_leave_type_widgets[k]["label"].setText(f"{new_val}%")

        self._popup_updating_sliders = False
        self._update_preview("leave_report")

    def _generate_leave_report_from_state(self, state_dict, escort_type: str) -> str:
        """Generate leave report text from a saved state dictionary."""
        p = self._get_pronouns()
        parts = []

        if not state_dict:
            return ""

        leaves_idx = state_dict.get('leaves', 0)
        frequency_idx = state_dict.get('frequency', 0)
        duration_idx = state_dict.get('duration', 0)

        leaves_options = ["1", "2", "3", "4", "5", "6", "7"]
        frequency_options = ["weekly", "2 weekly", "3 weekly", "monthly", "2 monthly"]
        duration_options = ["30 mins", "1 hour", "2 hours", "3 hours", "4 hours", "5 hours", "6 hours", "7 hours", "8 hours"]

        leaves = leaves_options[leaves_idx] if leaves_idx < len(leaves_options) else "1"
        frequency = frequency_options[frequency_idx] if frequency_idx < len(frequency_options) else "weekly"
        duration = duration_options[duration_idx] if duration_idx < len(duration_options) else "1 hour"

        leave_type_labels = {
            "ground": "ground leave",
            "local": "local community leave",
            "community": "community leave",
            "extended": "extended community leave",
            "overnight": "overnight leave"
        }

        leave_types = state_dict.get('leave_types', {})
        checked_types = []
        for key, lt in leave_types.items():
            if lt.get('checked', False):
                weight = lt.get('value', 0)
                checked_types.append((key, weight, leave_type_labels.get(key, key)))

        checked_types.sort(key=lambda x: x[1], reverse=True)

        if checked_types:
            type_phrases = []
            for i, (key, weight, label) in enumerate(checked_types):
                if i == 0:
                    type_phrases.append(f"mainly {label}")
                elif i == 1:
                    type_phrases.append(f"but also some {label}")
                elif i == 2:
                    type_phrases.append(f"and to a lesser extent {label}")
                else:
                    type_phrases.append(f"and occasionally {label}")

            type_str = ", ".join(type_phrases) if len(type_phrases) > 1 else type_phrases[0]
            parts.append(f"Over the past year, {p['subj_l']} {p['has']} taken approximately {leaves} {escort_type} leave{'s' if leaves != '1' else ''} {frequency}, averaging {duration} per leave, engaging in {type_str}.")

        # Other leave types
        other_types = []
        if state_dict.get('medical', False):
            other_types.append("medical appointments")
        if state_dict.get('court', False):
            other_types.append("court appearances")
        if state_dict.get('compassionate', False):
            other_types.append("compassionate visits")

        if other_types:
            if len(other_types) == 1:
                parts.append(f"{p['subj']} {p['has']} also taken leave for {other_types[0]}.")
            else:
                parts.append(f"{p['subj']} {p['has']} also taken leave for {', '.join(other_types[:-1])} and {other_types[-1]}.")

        # Leave suspension
        if state_dict.get('suspended_yes', False):
            details = state_dict.get('suspension_details', '').strip()
            if details:
                parts.append(f"{p['pos']} leave has been suspended in the past. {details}")
            else:
                parts.append(f"{p['pos']} leave has been suspended in the past.")
        elif state_dict.get('suspended_no', False):
            parts.append(f"{p['pos']} leave has never been suspended.")

        return " ".join(parts) if parts else ""

    def _generate_leave_report(self) -> str:
        """Generate leave report text from the 3g-style popup for both escorted and unescorted."""
        # Initialize state dicts if they don't exist
        if not hasattr(self, '_leave_escorted_state'):
            self._leave_escorted_state = {}
        if not hasattr(self, '_leave_unescorted_state'):
            self._leave_unescorted_state = {}

        # Save current state to the appropriate dictionary
        is_escorted = hasattr(self, 'popup_leave_escorted') and self.popup_leave_escorted.isChecked()
        if hasattr(self, 'popup_leaves_dropdown'):  # Only save if popup has been built
            if is_escorted:
                self._save_leave_state(self._leave_escorted_state)
            else:
                self._save_leave_state(self._leave_unescorted_state)

        # Generate both outputs
        escorted_text = self._generate_leave_report_from_state(self._leave_escorted_state, "escorted")
        unescorted_text = self._generate_leave_report_from_state(self._leave_unescorted_state, "unescorted")

        result_parts = []
        if escorted_text:
            result_parts.append(f"ESCORTED LEAVE: {escorted_text}")
        if unescorted_text:
            result_parts.append(f"UNESCORTED LEAVE: {unescorted_text}")

        result = "\n\n".join(result_parts) if result_parts else ""

        # Imported notes (checked entries)
        imported_texts = []
        if hasattr(self, 'popup_leave_report_imported_entries'):
            for entry in self.popup_leave_report_imported_entries:
                if entry["checkbox"].isChecked():
                    date = entry.get("date", "")
                    text = entry["text"]
                    if date:
                        imported_texts.append(f"[{date}] {text}")
                    else:
                        imported_texts.append(text)

        if imported_texts:
            if result:
                result += "\n\n--- Imported Notes ---\n" + "\n".join(imported_texts)
            else:
                result = "--- Imported Notes ---\n" + "\n".join(imported_texts)

        return result if result else "No leave information provided."

    def populate_popup_leave_report_imports(self, entries: list):
        """Populate the imported data panel in popup 14 (Leave Report) with leave-related entries.

        Searches for leave references in clinical notes.
        Filters to 12-month window from latest entry.
        """
        from datetime import datetime, timedelta
        import re

        print(f"[MOJ-ASR] Section 14 popup: populate_popup_leave_report_imports called with {len(entries) if entries else 0} entries")

        import_layout = getattr(self, 'popup_leave_report_import_layout', None)
        if not import_layout:
            print("[MOJ-ASR] Section 14 popup: popup_leave_report_import_layout not available")
            return

        # Clear existing entries
        while import_layout.count():
            item = import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.popup_leave_report_imported_entries = []

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        # Leave-related keywords
        LEAVE_KEYWORDS = [
            "leave", "escorted", "unescorted", "ground leave", "community leave",
            "local leave", "extended leave", "overnight leave", "section 17",
            "s17", "s.17", "leave of absence", "trial leave", "compassionate",
            "medical appointment", "court attendance", "accompanied", "unaccompanied",
            "leave suspended", "suspension", "awol", "absconded", "failed to return",
            "leave cancelled", "leave granted", "leave request", "leave application"
        ]

        parsed_entries = []
        latest_date = None

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date_val = entry.get("date") or entry.get("datetime")
            if not text:
                continue

            entry_date = None
            if isinstance(date_val, datetime):
                entry_date = date_val
            elif isinstance(date_val, str) and date_val:
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y %H:%M"]:
                    try:
                        entry_date = datetime.strptime(date_val.split()[0] if ' ' in date_val else date_val, fmt)
                        break
                    except:
                        pass

            parsed_entries.append({"text": text, "date": date_val, "date_obj": entry_date})
            if entry_date and (latest_date is None or entry_date > latest_date):
                latest_date = entry_date

        if latest_date:
            cutoff_date = latest_date - timedelta(days=365)
        else:
            cutoff_date = datetime.now() - timedelta(days=365)

        filtered_entries = []
        for entry in parsed_entries:
            entry_date = entry.get("date_obj")
            if entry_date is None or entry_date >= cutoff_date:
                filtered_entries.append(entry)

        if not filtered_entries:
            placeholder = QLabel("No notes from the last 12 months found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def find_leave_keywords(text):
            text_lower = text.lower()
            matches = []
            for kw in LEAVE_KEYWORDS:
                if kw in text_lower:
                    matches.append(kw)
            return matches

        categorized = []
        for entry in filtered_entries:
            matched = find_leave_keywords(entry["text"])
            if matched:
                snippet = entry["text"][:200] + "..." if len(entry["text"]) > 200 else entry["text"]
                entry["snippet"] = snippet
                entry["matched_keywords"] = matched
                categorized.append(entry)

        categorized.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)

        # Add category label to header showing count of Leave entries found
        category_buttons_layout = getattr(self, 'popup_leave_report_category_buttons_layout', None)
        if category_buttons_layout:
            while category_buttons_layout.count():
                item = category_buttons_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            if categorized:
                found_lbl = QLabel("Found:")
                found_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #806000; background: transparent;")
                category_buttons_layout.addWidget(found_lbl)
                lbl = QLabel(f"Leave ({len(categorized)})")
                lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: white; background: #6366f1; padding: 2px 6px; border-radius: 4px;")
                category_buttons_layout.addWidget(lbl)

        if not categorized:
            placeholder = QLabel("No leave-related notes found.")
            placeholder.setStyleSheet("color: #806000; font-style: italic; font-size: 15px; background: transparent;")
            import_layout.addWidget(placeholder)
            return

        def format_date_nice(date_obj):
            if not date_obj:
                return ""
            if isinstance(date_obj, str):
                return date_obj
            day = date_obj.day
            suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
            return date_obj.strftime(f"{day}{suffix} %b %Y")

        def hex_to_light_bg(hex_color):
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = int(r * 0.3 + 255 * 0.7)
            g = int(g * 0.3 + 255 * 0.7)
            b = int(b * 0.3 + 255 * 0.7)
            return f"#{r:02x}{g:02x}{b:02x}"

        def highlight_keywords(text, matched_keywords):
            import html
            escaped = html.escape(text)
            leave_color = "#6366f1"  # Indigo for leave
            light_color = hex_to_light_bg(leave_color)

            sorted_keywords = sorted(matched_keywords, key=len, reverse=True)
            for kw in sorted_keywords:
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                escaped = pattern.sub(
                    lambda m: f'<span style="background-color: {light_color}; font-weight: 600;">{m.group()}</span>',
                    escaped
                )
            escaped = escaped.replace('\n', '<br>')
            return escaped

        for entry in categorized:
            full_text = '\n'.join(line for line in entry["text"].split('\n') if line.strip())
            date_obj = entry.get("date_obj")
            date_raw = entry.get("date", "")
            matched_keywords = entry.get("matched_keywords", [])
            date_str = format_date_nice(date_obj) if date_obj else str(date_raw) if date_raw else ""
            highlighted_html = highlight_keywords(full_text, matched_keywords)

            entry_frame, cb, body_text = self._create_4a_import_entry(
                text=full_text,
                date_str=date_str,
                preview_key="leave_report",
                categories=["Leave"],
                category_colors={"Leave": "#6366f1"},
                highlighted_html=highlighted_html,
                parent_container=getattr(self, 'popup_leave_report_import_container', None),
            )

            import_layout.addWidget(entry_frame)
            self.popup_leave_report_imported_entries.append({"checkbox": cb, "text": full_text, "date": date_raw, "frame": entry_frame})

        print(f"[MOJ-ASR] Section 14 popup: Added {len(self.popup_leave_report_imported_entries)} leave entries")

    def _build_popup_additional_comments(self):
        """Build additional comments popup with legal criteria section - matches GPRLegalCriteriaPopup."""
        container, layout = self._create_popup_container("additional_comments")

        # ============================================
        # LEGAL CRITERIA SECTION
        # ============================================
        legal_frame = QFrame()
        legal_frame.setStyleSheet("""
            QFrame { background: #f0fdf4; border: 1px solid #86efac; border-radius: 8px; }
            QLabel { background: transparent; border: none; }
            QRadioButton, QCheckBox { background: transparent; border: none; }
        """)
        legal_layout = QVBoxLayout(legal_frame)
        legal_layout.setContentsMargins(12, 10, 12, 10)
        legal_layout.setSpacing(8)

        legal_header = QLabel("<b>Legal Criteria for Continued Detention</b>")
        legal_header.setStyleSheet("font-size: 18px; color: #166534;")
        legal_layout.addWidget(legal_header)

        # ============================================
        # 1. MENTAL DISORDER - Present/Absent
        # ============================================
        md_label = QLabel("Mental Disorder")
        md_label.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        legal_layout.addWidget(md_label)

        self.popup_md_group = QButtonGroup(self)
        md_row = QHBoxLayout()
        md_row.setSpacing(16)
        radio_style = "QRadioButton { font-size: 18px; }"
        self.popup_md_present = QRadioButton("Present")
        self.popup_md_present.setStyleSheet(radio_style)
        self.popup_md_absent = QRadioButton("Absent")
        self.popup_md_absent.setStyleSheet(radio_style)
        self.popup_md_group.addButton(self.popup_md_present, 0)
        self.popup_md_group.addButton(self.popup_md_absent, 1)
        md_row.addWidget(self.popup_md_present)
        md_row.addWidget(self.popup_md_absent)
        md_row.addStretch()
        legal_layout.addLayout(md_row)

        # ICD-10 dropdown container (shown when Present) - using grouped categories like Section 3
        self.popup_icd10_container = QWidget()
        self.popup_icd10_container.setStyleSheet("background: transparent;")
        icd10_layout = QHBoxLayout(self.popup_icd10_container)
        icd10_layout.setContentsMargins(0, 4, 0, 4)
        icd10_layout.setSpacing(8)
        icd10_label = QLabel("ICD-10 Diagnosis:")
        icd10_label.setStyleSheet("font-size: 16px; color: #6b7280;")
        icd10_layout.addWidget(icd10_label)
        self.popup_icd10_combo = QComboBox()
        self.popup_icd10_combo.setStyleSheet("font-size: 16px; padding: 4px;")
        self.popup_icd10_combo.setMinimumWidth(400)
        self.popup_icd10_combo.setEditable(True)

        # Add grouped ICD-10 options (reuse ICD10_GROUPED from Section 3 if available)
        icd10_grouped = getattr(self, 'ICD10_GROUPED', None)
        if not icd10_grouped:
            icd10_grouped = [
                ("Schizophrenia & Psychosis", [
                    ("F20.0", "Paranoid schizophrenia"),
                    ("F20.1", "Hebephrenic schizophrenia"),
                    ("F20.2", "Catatonic schizophrenia"),
                    ("F20.3", "Undifferentiated schizophrenia"),
                    ("F20.5", "Residual schizophrenia"),
                    ("F20.6", "Simple schizophrenia"),
                    ("F20.9", "Schizophrenia, unspecified"),
                    ("F21", "Schizotypal disorder"),
                    ("F22", "Persistent delusional disorders"),
                    ("F23", "Acute and transient psychotic disorders"),
                    ("F25.0", "Schizoaffective disorder, manic type"),
                    ("F25.1", "Schizoaffective disorder, depressive type"),
                    ("F25.2", "Schizoaffective disorder, mixed type"),
                    ("F25.9", "Schizoaffective disorder, unspecified"),
                    ("F29", "Unspecified nonorganic psychosis"),
                ]),
                ("Mood Disorders - Bipolar", [
                    ("F30.0", "Hypomania"),
                    ("F30.1", "Mania without psychotic symptoms"),
                    ("F30.2", "Mania with psychotic symptoms"),
                    ("F31.0", "Bipolar disorder, current episode hypomanic"),
                    ("F31.1", "Bipolar disorder, current episode manic without psychosis"),
                    ("F31.2", "Bipolar disorder, current episode manic with psychosis"),
                    ("F31.3", "Bipolar disorder, current episode mild/moderate depression"),
                    ("F31.4", "Bipolar disorder, current episode severe depression without psychosis"),
                    ("F31.5", "Bipolar disorder, current episode severe depression with psychosis"),
                    ("F31.6", "Bipolar disorder, current episode mixed"),
                    ("F31.7", "Bipolar disorder, currently in remission"),
                    ("F31.9", "Bipolar disorder, unspecified"),
                ]),
                ("Mood Disorders - Depression", [
                    ("F32.0", "Mild depressive episode"),
                    ("F32.1", "Moderate depressive episode"),
                    ("F32.2", "Severe depressive episode without psychosis"),
                    ("F32.3", "Severe depressive episode with psychosis"),
                    ("F32.9", "Depressive episode, unspecified"),
                    ("F33.0", "Recurrent depression, current episode mild"),
                    ("F33.1", "Recurrent depression, current episode moderate"),
                    ("F33.2", "Recurrent depression, current episode severe without psychosis"),
                    ("F33.3", "Recurrent depression, current episode severe with psychosis"),
                    ("F33.9", "Recurrent depressive disorder, unspecified"),
                ]),
                ("Anxiety Disorders", [
                    ("F40.0", "Agoraphobia"),
                    ("F40.1", "Social phobias"),
                    ("F40.2", "Specific (isolated) phobias"),
                    ("F41.0", "Panic disorder"),
                    ("F41.1", "Generalized anxiety disorder"),
                    ("F41.2", "Mixed anxiety and depressive disorder"),
                    ("F42", "Obsessive-compulsive disorder"),
                    ("F43.0", "Acute stress reaction"),
                    ("F43.1", "Post-traumatic stress disorder"),
                    ("F43.2", "Adjustment disorders"),
                ]),
                ("Eating Disorders", [
                    ("F50.0", "Anorexia nervosa"),
                    ("F50.2", "Bulimia nervosa"),
                ]),
                ("Personality Disorders", [
                    ("F60.0", "Paranoid personality disorder"),
                    ("F60.1", "Schizoid personality disorder"),
                    ("F60.2", "Dissocial personality disorder"),
                    ("F60.3", "Emotionally unstable personality disorder"),
                    ("F60.4", "Histrionic personality disorder"),
                    ("F60.5", "Anankastic personality disorder"),
                    ("F60.6", "Anxious personality disorder"),
                    ("F60.7", "Dependent personality disorder"),
                    ("F60.9", "Personality disorder, unspecified"),
                ]),
                ("Intellectual Disability", [
                    ("F70", "Mild intellectual disability"),
                    ("F71", "Moderate intellectual disability"),
                    ("F72", "Severe intellectual disability"),
                    ("F79", "Unspecified intellectual disability"),
                ]),
                ("Organic Disorders", [
                    ("F00", "Dementia in Alzheimer's disease"),
                    ("F01", "Vascular dementia"),
                    ("F03", "Unspecified dementia"),
                    ("F05", "Delirium"),
                    ("F06", "Other mental disorders due to brain damage"),
                ]),
                ("Substance Use Disorders", [
                    ("F10", "Mental disorders due to alcohol"),
                    ("F11", "Mental disorders due to opioids"),
                    ("F12", "Mental disorders due to cannabinoids"),
                    ("F14", "Mental disorders due to cocaine"),
                    ("F15", "Mental disorders due to stimulants"),
                    ("F19", "Mental disorders due to multiple drug use"),
                ]),
            ]

        self.popup_icd10_combo.addItem("Select...")
        all_items = ["Select..."]

        for group_name, diagnoses in icd10_grouped:
            # Add category header (disabled, styled)
            self.popup_icd10_combo.addItem(f"── {group_name} ──")
            idx = self.popup_icd10_combo.count() - 1
            self.popup_icd10_combo.model().item(idx).setEnabled(False)
            self.popup_icd10_combo.model().item(idx).setData(QColor("#6b7280"), Qt.ItemDataRole.ForegroundRole)
            font = self.popup_icd10_combo.model().item(idx).font()
            font.setBold(True)
            self.popup_icd10_combo.model().item(idx).setFont(font)

            # Add diagnoses
            for code, name in diagnoses:
                display_text = f"{code} {name}"
                self.popup_icd10_combo.addItem(display_text, code)
                all_items.append(display_text)

        # Enable autocomplete
        self.popup_icd10_combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        completer = QCompleter(all_items)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.popup_icd10_combo.setCompleter(completer)

        self.popup_icd10_combo.currentTextChanged.connect(lambda: self._update_preview("additional_comments"))
        icd10_layout.addWidget(self.popup_icd10_combo)
        icd10_layout.addStretch()
        self.popup_icd10_container.hide()
        legal_layout.addWidget(self.popup_icd10_container)

        # Criteria container (shown when Present)
        self.popup_criteria_container = QWidget()
        self.popup_criteria_container.setStyleSheet("background: transparent;")
        criteria_layout = QVBoxLayout(self.popup_criteria_container)
        criteria_layout.setContentsMargins(0, 8, 0, 0)
        criteria_layout.setSpacing(12)

        # ============================================
        # 2. CRITERIA WARRANTING DETENTION - Met/Not Met
        # ============================================
        cwd_label = QLabel("Criteria Warranting Detention")
        cwd_label.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        criteria_layout.addWidget(cwd_label)

        self.popup_cwd_group = QButtonGroup(self)
        cwd_row = QHBoxLayout()
        cwd_row.setSpacing(16)
        self.popup_cwd_met = QRadioButton("Met")
        self.popup_cwd_met.setStyleSheet(radio_style)
        self.popup_cwd_not_met = QRadioButton("Not Met")
        self.popup_cwd_not_met.setStyleSheet(radio_style)
        self.popup_cwd_group.addButton(self.popup_cwd_met, 0)
        self.popup_cwd_group.addButton(self.popup_cwd_not_met, 1)
        cwd_row.addWidget(self.popup_cwd_met)
        cwd_row.addWidget(self.popup_cwd_not_met)
        cwd_row.addStretch()
        criteria_layout.addLayout(cwd_row)

        # Nature/Degree container (shown when Met)
        self.popup_nd_container = QWidget()
        self.popup_nd_container.setStyleSheet("background: transparent;")
        nd_layout = QVBoxLayout(self.popup_nd_container)
        nd_layout.setContentsMargins(16, 8, 0, 0)
        nd_layout.setSpacing(8)

        # Nature checkbox
        self.popup_nature_cb = QCheckBox("Nature")
        self.popup_nature_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        nd_layout.addWidget(self.popup_nature_cb)

        # Nature sub-options container
        self.popup_nature_options_container = QWidget()
        self.popup_nature_options_container.setStyleSheet("background: transparent;")
        nature_opt_layout = QVBoxLayout(self.popup_nature_options_container)
        nature_opt_layout.setContentsMargins(24, 4, 0, 0)
        nature_opt_layout.setSpacing(4)

        self.popup_relapsing_cb = QCheckBox("Relapsing and remitting")
        self.popup_relapsing_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        nature_opt_layout.addWidget(self.popup_relapsing_cb)

        self.popup_treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.popup_treatment_resistant_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        nature_opt_layout.addWidget(self.popup_treatment_resistant_cb)

        self.popup_chronic_cb = QCheckBox("Chronic and enduring")
        self.popup_chronic_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        nature_opt_layout.addWidget(self.popup_chronic_cb)

        self.popup_nature_options_container.hide()
        nd_layout.addWidget(self.popup_nature_options_container)

        # Degree checkbox
        self.popup_degree_cb = QCheckBox("Degree")
        self.popup_degree_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        nd_layout.addWidget(self.popup_degree_cb)

        # Degree sub-options container
        self.popup_degree_options_container = QWidget()
        self.popup_degree_options_container.setStyleSheet("background: transparent;")
        degree_opt_layout = QVBoxLayout(self.popup_degree_options_container)
        degree_opt_layout.setContentsMargins(24, 4, 0, 0)
        degree_opt_layout.setSpacing(8)

        # Severity slider
        slider_label = QLabel("Symptom severity:")
        slider_label.setStyleSheet("font-size: 16px; color: #6b7280;")
        degree_opt_layout.addWidget(slider_label)

        slider_row = QHBoxLayout()
        slider_row.setSpacing(8)

        self.popup_degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.popup_degree_slider.setMinimum(1)
        self.popup_degree_slider.setMaximum(4)
        self.popup_degree_slider.setValue(2)
        self.popup_degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.popup_degree_slider.setTickInterval(1)
        self.popup_degree_slider.setFixedWidth(200)
        self.popup_degree_slider.setStyleSheet("""
            QSlider::groove:horizontal { background: #d1d5db; height: 4px; border-radius: 2px; }
            QSlider::handle:horizontal { background: #166534; width: 12px; margin: -4px 0; border-radius: 6px; }
        """)
        self.popup_degree_slider.valueChanged.connect(self._on_legal_degree_slider_changed)
        slider_row.addWidget(self.popup_degree_slider)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.popup_degree_level = QLabel("Several")
        self.popup_degree_level.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500;")
        degree_opt_layout.addWidget(self.popup_degree_level)

        # Details text box
        details_label = QLabel("Symptoms including:")
        details_label.setStyleSheet("font-size: 16px; color: #6b7280;")
        degree_opt_layout.addWidget(details_label)

        self.popup_degree_details = QTextEdit()
        self.popup_degree_details.setPlaceholderText("Enter symptom details...")
        self.popup_degree_details.setMaximumHeight(80)
        self.popup_degree_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 6px; padding: 6px; background: white; font-size: 17px;")
        self.popup_degree_details.textChanged.connect(lambda: self._update_preview("additional_comments"))
        degree_opt_layout.addWidget(self.popup_degree_details)

        self.popup_degree_options_container.hide()
        nd_layout.addWidget(self.popup_degree_options_container)

        self.popup_nd_container.hide()
        criteria_layout.addWidget(self.popup_nd_container)

        # ============================================
        # 3. NECESSITY SECTION
        # ============================================
        necessity_label = QLabel("Necessity")
        necessity_label.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        criteria_layout.addWidget(necessity_label)

        self.popup_nec_group = QButtonGroup(self)
        nec_row = QHBoxLayout()
        nec_row.setSpacing(16)
        self.popup_nec_yes = QRadioButton("Yes")
        self.popup_nec_yes.setStyleSheet(radio_style)
        self.popup_nec_no = QRadioButton("No")
        self.popup_nec_no.setStyleSheet(radio_style)
        self.popup_nec_group.addButton(self.popup_nec_yes, 0)
        self.popup_nec_group.addButton(self.popup_nec_no, 1)
        nec_row.addWidget(self.popup_nec_yes)
        nec_row.addWidget(self.popup_nec_no)
        nec_row.addStretch()
        criteria_layout.addLayout(nec_row)

        # Health & Safety container
        self.popup_hs_container = QWidget()
        self.popup_hs_container.setStyleSheet("background: transparent;")
        hs_layout = QVBoxLayout(self.popup_hs_container)
        hs_layout.setContentsMargins(16, 8, 0, 0)
        hs_layout.setSpacing(8)

        # Health checkbox
        self.popup_health_cb = QCheckBox("Health")
        self.popup_health_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        hs_layout.addWidget(self.popup_health_cb)

        # Health sub-options container
        self.popup_health_container = QWidget()
        self.popup_health_container.setStyleSheet("background: transparent;")
        health_layout = QVBoxLayout(self.popup_health_container)
        health_layout.setContentsMargins(24, 4, 0, 0)
        health_layout.setSpacing(4)

        # Mental Health checkbox
        self.popup_mental_health_cb = QCheckBox("Mental Health")
        self.popup_mental_health_cb.toggled.connect(self._on_popup_mental_health_toggled)
        health_layout.addWidget(self.popup_mental_health_cb)

        # Mental Health sub-options
        self.popup_mental_health_container = QWidget()
        self.popup_mental_health_container.setStyleSheet("background: transparent;")
        mh_layout = QVBoxLayout(self.popup_mental_health_container)
        mh_layout.setContentsMargins(24, 4, 0, 0)
        mh_layout.setSpacing(4)

        self.popup_poor_compliance_cb = QCheckBox("Poor compliance")
        self.popup_poor_compliance_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        mh_layout.addWidget(self.popup_poor_compliance_cb)

        self.popup_limited_insight_cb = QCheckBox("Limited insight")
        self.popup_limited_insight_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        mh_layout.addWidget(self.popup_limited_insight_cb)

        self.popup_mental_health_container.hide()
        health_layout.addWidget(self.popup_mental_health_container)

        # Physical Health checkbox
        self.popup_physical_health_cb = QCheckBox("Physical Health")
        self.popup_physical_health_cb.toggled.connect(self._on_popup_physical_health_toggled)
        health_layout.addWidget(self.popup_physical_health_cb)

        # Physical Health details
        self.popup_physical_health_details = QTextEdit()
        self.popup_physical_health_details.setPlaceholderText("Enter physical health details...")
        self.popup_physical_health_details.setMaximumHeight(60)
        self.popup_physical_health_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 4px; background: white; font-size: 17px;")
        self.popup_physical_health_details.textChanged.connect(lambda: self._update_preview("additional_comments"))
        self.popup_physical_health_details.hide()
        health_layout.addWidget(self.popup_physical_health_details)

        self.popup_health_container.hide()
        hs_layout.addWidget(self.popup_health_container)

        # Safety checkbox
        self.popup_safety_cb = QCheckBox("Safety")
        self.popup_safety_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        hs_layout.addWidget(self.popup_safety_cb)

        # Safety sub-options container
        self.popup_safety_container = QWidget()
        self.popup_safety_container.setStyleSheet("background: transparent;")
        safety_layout = QVBoxLayout(self.popup_safety_container)
        safety_layout.setContentsMargins(24, 4, 0, 0)
        safety_layout.setSpacing(4)

        # Self checkbox
        self.popup_self_cb = QCheckBox("Self")
        self.popup_self_cb.toggled.connect(self._on_popup_self_toggled)
        safety_layout.addWidget(self.popup_self_cb)

        # Self options container with checkboxes (Historical/Current columns)
        self.popup_self_options_container = QWidget()
        self.popup_self_options_container.setStyleSheet("background: transparent;")
        self_opt_layout = QVBoxLayout(self.popup_self_options_container)
        self_opt_layout.setContentsMargins(24, 4, 0, 0)
        self_opt_layout.setSpacing(4)

        # Header row for Self
        self_header = QHBoxLayout()
        self_header.setSpacing(8)
        self_spacer_lbl = QLabel("")
        self_spacer_lbl.setFixedWidth(120)  # Match label column width
        self_header.addWidget(self_spacer_lbl)
        self_hist_lbl = QLabel("Historical")
        self_hist_lbl.setStyleSheet("font-size: 14px; color: #6b7280; font-weight: 600;")
        self_hist_lbl.setFixedWidth(70)
        self_header.addWidget(self_hist_lbl)
        self_curr_lbl = QLabel("Current")
        self_curr_lbl.setStyleSheet("font-size: 14px; color: #6b7280; font-weight: 600;")
        self_curr_lbl.setFixedWidth(70)
        self_header.addWidget(self_curr_lbl)
        self_header.addStretch()
        self_opt_layout.addLayout(self_header)

        # Self-neglect row
        self_neglect_row = QHBoxLayout()
        self_neglect_row.setSpacing(8)
        self_neglect_lbl = QLabel("Self-neglect")
        self_neglect_lbl.setStyleSheet("font-size: 15px;")
        self_neglect_lbl.setFixedWidth(120)
        self_neglect_row.addWidget(self_neglect_lbl)
        self.popup_self_neglect_hist = QCheckBox()
        self.popup_self_neglect_hist.setFixedWidth(70)
        self.popup_self_neglect_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        self_neglect_row.addWidget(self.popup_self_neglect_hist)
        self.popup_self_neglect_curr = QCheckBox()
        self.popup_self_neglect_curr.setFixedWidth(70)
        self.popup_self_neglect_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        self_neglect_row.addWidget(self.popup_self_neglect_curr)
        self_neglect_row.addStretch()
        self_opt_layout.addLayout(self_neglect_row)

        # Risky behaviour row
        self_risky_row = QHBoxLayout()
        self_risky_row.setSpacing(8)
        self_risky_lbl = QLabel("Risky behaviour")
        self_risky_lbl.setStyleSheet("font-size: 15px;")
        self_risky_lbl.setFixedWidth(120)
        self_risky_row.addWidget(self_risky_lbl)
        self.popup_self_risky_hist = QCheckBox()
        self.popup_self_risky_hist.setFixedWidth(70)
        self.popup_self_risky_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        self_risky_row.addWidget(self.popup_self_risky_hist)
        self.popup_self_risky_curr = QCheckBox()
        self.popup_self_risky_curr.setFixedWidth(70)
        self.popup_self_risky_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        self_risky_row.addWidget(self.popup_self_risky_curr)
        self_risky_row.addStretch()
        self_opt_layout.addLayout(self_risky_row)

        # Self-harm row
        self_harm_row = QHBoxLayout()
        self_harm_row.setSpacing(8)
        self_harm_lbl = QLabel("Self-harm")
        self_harm_lbl.setStyleSheet("font-size: 15px;")
        self_harm_lbl.setFixedWidth(120)
        self_harm_row.addWidget(self_harm_lbl)
        self.popup_self_harm_hist = QCheckBox()
        self.popup_self_harm_hist.setFixedWidth(70)
        self.popup_self_harm_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        self_harm_row.addWidget(self.popup_self_harm_hist)
        self.popup_self_harm_curr = QCheckBox()
        self.popup_self_harm_curr.setFixedWidth(70)
        self.popup_self_harm_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        self_harm_row.addWidget(self.popup_self_harm_curr)
        self_harm_row.addStretch()
        self_opt_layout.addLayout(self_harm_row)

        # Self details text field
        self.popup_self_details = QTextEdit()
        self.popup_self_details.setPlaceholderText("Enter additional details about risk to self...")
        self.popup_self_details.setMaximumHeight(50)
        self.popup_self_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 4px; background: white; font-size: 15px;")
        self.popup_self_details.textChanged.connect(lambda: self._update_preview("additional_comments"))
        self_opt_layout.addWidget(self.popup_self_details)

        self.popup_self_options_container.hide()
        safety_layout.addWidget(self.popup_self_options_container)

        # Others checkbox
        self.popup_others_cb = QCheckBox("Others")
        self.popup_others_cb.toggled.connect(self._on_popup_others_toggled)
        safety_layout.addWidget(self.popup_others_cb)

        # Others options container with checkboxes (Historical/Current columns)
        self.popup_others_options_container = QWidget()
        self.popup_others_options_container.setStyleSheet("background: transparent;")
        others_opt_layout = QVBoxLayout(self.popup_others_options_container)
        others_opt_layout.setContentsMargins(24, 4, 0, 0)
        others_opt_layout.setSpacing(4)

        # Header row for Others
        others_header = QHBoxLayout()
        others_header.setSpacing(8)
        others_spacer_lbl = QLabel("")
        others_spacer_lbl.setFixedWidth(120)  # Match label column width
        others_header.addWidget(others_spacer_lbl)
        others_hist_lbl = QLabel("Historical")
        others_hist_lbl.setStyleSheet("font-size: 14px; color: #6b7280; font-weight: 600;")
        others_hist_lbl.setFixedWidth(70)
        others_header.addWidget(others_hist_lbl)
        others_curr_lbl = QLabel("Current")
        others_curr_lbl.setStyleSheet("font-size: 14px; color: #6b7280; font-weight: 600;")
        others_curr_lbl.setFixedWidth(70)
        others_header.addWidget(others_curr_lbl)
        others_header.addStretch()
        others_opt_layout.addLayout(others_header)

        # Violence row
        violence_row = QHBoxLayout()
        violence_row.setSpacing(8)
        violence_lbl = QLabel("Violence")
        violence_lbl.setStyleSheet("font-size: 15px;")
        violence_lbl.setFixedWidth(120)
        violence_row.addWidget(violence_lbl)
        self.popup_violence_hist = QCheckBox()
        self.popup_violence_hist.setFixedWidth(70)
        self.popup_violence_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        violence_row.addWidget(self.popup_violence_hist)
        self.popup_violence_curr = QCheckBox()
        self.popup_violence_curr.setFixedWidth(70)
        self.popup_violence_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        violence_row.addWidget(self.popup_violence_curr)
        violence_row.addStretch()
        others_opt_layout.addLayout(violence_row)

        # Verbal aggression row
        verbal_row = QHBoxLayout()
        verbal_row.setSpacing(8)
        verbal_lbl = QLabel("Verbal aggression")
        verbal_lbl.setStyleSheet("font-size: 15px;")
        verbal_lbl.setFixedWidth(120)
        verbal_row.addWidget(verbal_lbl)
        self.popup_verbal_hist = QCheckBox()
        self.popup_verbal_hist.setFixedWidth(70)
        self.popup_verbal_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        verbal_row.addWidget(self.popup_verbal_hist)
        self.popup_verbal_curr = QCheckBox()
        self.popup_verbal_curr.setFixedWidth(70)
        self.popup_verbal_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        verbal_row.addWidget(self.popup_verbal_curr)
        verbal_row.addStretch()
        others_opt_layout.addLayout(verbal_row)

        # Sexual violence row
        sexual_row = QHBoxLayout()
        sexual_row.setSpacing(8)
        sexual_lbl = QLabel("Sexual violence")
        sexual_lbl.setStyleSheet("font-size: 15px;")
        sexual_lbl.setFixedWidth(120)
        sexual_row.addWidget(sexual_lbl)
        self.popup_sexual_hist = QCheckBox()
        self.popup_sexual_hist.setFixedWidth(70)
        self.popup_sexual_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        sexual_row.addWidget(self.popup_sexual_hist)
        self.popup_sexual_curr = QCheckBox()
        self.popup_sexual_curr.setFixedWidth(70)
        self.popup_sexual_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        sexual_row.addWidget(self.popup_sexual_curr)
        sexual_row.addStretch()
        others_opt_layout.addLayout(sexual_row)

        # Stalking row
        stalking_row = QHBoxLayout()
        stalking_row.setSpacing(8)
        stalking_lbl = QLabel("Stalking")
        stalking_lbl.setStyleSheet("font-size: 15px;")
        stalking_lbl.setFixedWidth(120)
        stalking_row.addWidget(stalking_lbl)
        self.popup_stalking_hist = QCheckBox()
        self.popup_stalking_hist.setFixedWidth(70)
        self.popup_stalking_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        stalking_row.addWidget(self.popup_stalking_hist)
        self.popup_stalking_curr = QCheckBox()
        self.popup_stalking_curr.setFixedWidth(70)
        self.popup_stalking_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        stalking_row.addWidget(self.popup_stalking_curr)
        stalking_row.addStretch()
        others_opt_layout.addLayout(stalking_row)

        # Arson row
        arson_row = QHBoxLayout()
        arson_row.setSpacing(8)
        arson_lbl = QLabel("Arson")
        arson_lbl.setStyleSheet("font-size: 15px;")
        arson_lbl.setFixedWidth(120)
        arson_row.addWidget(arson_lbl)
        self.popup_arson_hist = QCheckBox()
        self.popup_arson_hist.setFixedWidth(70)
        self.popup_arson_hist.toggled.connect(lambda: self._update_preview("additional_comments"))
        arson_row.addWidget(self.popup_arson_hist)
        self.popup_arson_curr = QCheckBox()
        self.popup_arson_curr.setFixedWidth(70)
        self.popup_arson_curr.toggled.connect(lambda: self._update_preview("additional_comments"))
        arson_row.addWidget(self.popup_arson_curr)
        arson_row.addStretch()
        others_opt_layout.addLayout(arson_row)

        # Others details text field
        self.popup_others_details = QTextEdit()
        self.popup_others_details.setPlaceholderText("Enter additional details about risk to others...")
        self.popup_others_details.setMaximumHeight(50)
        self.popup_others_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 4px; background: white; font-size: 15px;")
        self.popup_others_details.textChanged.connect(lambda: self._update_preview("additional_comments"))
        others_opt_layout.addWidget(self.popup_others_details)

        self.popup_others_options_container.hide()
        safety_layout.addWidget(self.popup_others_options_container)

        self.popup_safety_container.hide()
        hs_layout.addWidget(self.popup_safety_container)

        self.popup_hs_container.hide()
        criteria_layout.addWidget(self.popup_hs_container)

        # ============================================
        # 4. TREATMENT AVAILABLE
        # ============================================
        self.popup_treatment_cb = QCheckBox("Treatment Available")
        self.popup_treatment_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        self.popup_treatment_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        criteria_layout.addWidget(self.popup_treatment_cb)

        # ============================================
        # 5. LEAST RESTRICTIVE
        # ============================================
        self.popup_least_restrictive_cb = QCheckBox("Least Restrictive Option")
        self.popup_least_restrictive_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.popup_least_restrictive_cb.toggled.connect(lambda: self._update_preview("additional_comments"))
        criteria_layout.addWidget(self.popup_least_restrictive_cb)

        self.popup_criteria_container.hide()
        legal_layout.addWidget(self.popup_criteria_container)

        # ============================================
        # WIRE UP VISIBILITY TOGGLING
        # ============================================
        def on_md_changed():
            if self.popup_md_present.isChecked():
                self.popup_criteria_container.show()
                self.popup_icd10_container.show()
            else:
                self.popup_criteria_container.hide()
                self.popup_icd10_container.hide()
            self._update_preview("additional_comments")

        def on_cwd_changed():
            if self.popup_cwd_met.isChecked():
                self.popup_nd_container.show()
            else:
                self.popup_nd_container.hide()
                self.popup_nature_cb.setChecked(False)
                self.popup_degree_cb.setChecked(False)
            self._update_preview("additional_comments")

        def on_nature_toggled(checked):
            self.popup_nature_options_container.setVisible(checked)
            if not checked:
                self.popup_relapsing_cb.setChecked(False)
                self.popup_treatment_resistant_cb.setChecked(False)
                self.popup_chronic_cb.setChecked(False)
            self._update_preview("additional_comments")

        def on_degree_toggled(checked):
            self.popup_degree_options_container.setVisible(checked)
            if not checked:
                self.popup_degree_details.clear()
            self._update_preview("additional_comments")

        def on_nec_changed():
            if self.popup_nec_yes.isChecked():
                self.popup_hs_container.show()
            else:
                self.popup_hs_container.hide()
                self.popup_health_cb.setChecked(False)
                self.popup_safety_cb.setChecked(False)
            self._update_preview("additional_comments")

        def on_health_toggled(checked):
            self.popup_health_container.setVisible(checked)
            if not checked:
                self.popup_mental_health_cb.setChecked(False)
                self.popup_physical_health_cb.setChecked(False)
            self._update_preview("additional_comments")

        def on_safety_toggled(checked):
            self.popup_safety_container.setVisible(checked)
            if not checked:
                self.popup_self_cb.setChecked(False)
                self.popup_others_cb.setChecked(False)
            self._update_preview("additional_comments")

        self.popup_md_present.toggled.connect(on_md_changed)
        self.popup_md_absent.toggled.connect(on_md_changed)
        self.popup_cwd_met.toggled.connect(on_cwd_changed)
        self.popup_cwd_not_met.toggled.connect(on_cwd_changed)
        self.popup_nature_cb.toggled.connect(on_nature_toggled)
        self.popup_degree_cb.toggled.connect(on_degree_toggled)
        self.popup_nec_yes.toggled.connect(on_nec_changed)
        self.popup_nec_no.toggled.connect(on_nec_changed)
        self.popup_health_cb.toggled.connect(on_health_toggled)
        self.popup_safety_cb.toggled.connect(on_safety_toggled)

        layout.addWidget(legal_frame)

        # ============================================
        # ADDITIONAL COMMENTS TEXT
        # ============================================
        layout.addSpacing(12)
        layout.addWidget(QLabel("Additional comments:"))
        self.popup_additional = QTextEdit()
        self.popup_additional.setMaximumHeight(100)
        self.popup_additional.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px;")
        layout.addWidget(self.popup_additional)

        layout.addStretch()
        self._add_send_button(layout, "additional_comments", self._generate_additional_comments)
        self._connect_preview_updates("additional_comments", [self.popup_additional])

    def _on_popup_mental_health_toggled(self, checked):
        """Toggle Mental Health sub-options visibility."""
        self.popup_mental_health_container.setVisible(checked)
        if not checked:
            self.popup_poor_compliance_cb.setChecked(False)
            self.popup_limited_insight_cb.setChecked(False)
        self._update_preview("additional_comments")

    def _on_popup_physical_health_toggled(self, checked):
        """Toggle Physical Health details visibility."""
        self.popup_physical_health_details.setVisible(checked)
        if not checked:
            self.popup_physical_health_details.clear()
        self._update_preview("additional_comments")

    def _on_popup_self_toggled(self, checked):
        """Toggle Self options visibility."""
        self.popup_self_options_container.setVisible(checked)
        if not checked:
            # Clear all self checkboxes
            self.popup_self_neglect_hist.setChecked(False)
            self.popup_self_neglect_curr.setChecked(False)
            self.popup_self_risky_hist.setChecked(False)
            self.popup_self_risky_curr.setChecked(False)
            self.popup_self_harm_hist.setChecked(False)
            self.popup_self_harm_curr.setChecked(False)
            self.popup_self_details.clear()
        self._update_preview("additional_comments")

    def _on_popup_others_toggled(self, checked):
        """Toggle Others options visibility."""
        self.popup_others_options_container.setVisible(checked)
        if not checked:
            # Clear all others checkboxes
            self.popup_violence_hist.setChecked(False)
            self.popup_violence_curr.setChecked(False)
            self.popup_verbal_hist.setChecked(False)
            self.popup_verbal_curr.setChecked(False)
            self.popup_sexual_hist.setChecked(False)
            self.popup_sexual_curr.setChecked(False)
            self.popup_stalking_hist.setChecked(False)
            self.popup_stalking_curr.setChecked(False)
            self.popup_arson_hist.setChecked(False)
            self.popup_arson_curr.setChecked(False)
            self.popup_others_details.clear()
        self._update_preview("additional_comments")

    def _on_legal_degree_slider_changed(self, value):
        """Update degree level label based on slider value."""
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        if hasattr(self, 'popup_degree_level'):
            self.popup_degree_level.setText(levels.get(value, "Several"))
        self._update_preview("additional_comments")

    def _generate_additional_comments(self) -> str:
        """Generate additional comments text including legal criteria - flowing paragraph style."""
        p = self._get_pronouns()
        parts = []

        # 1. Mental Disorder + ICD-10 + Nature/Degree
        if hasattr(self, 'popup_md_present') and self.popup_md_present.isChecked():
            # Check for ICD-10 diagnosis
            icd10_text = ""
            if hasattr(self, 'popup_icd10_combo'):
                icd10_selection = self.popup_icd10_combo.currentText()
                if icd10_selection and icd10_selection != "Select...":
                    icd10_text = f" ({icd10_selection})"

            md_base = f"{p['subj']} {p['suffers']} from a mental disorder{icd10_text} under the Mental Health Act"

            if hasattr(self, 'popup_cwd_met') and self.popup_cwd_met.isChecked():
                nature_checked = hasattr(self, 'popup_nature_cb') and self.popup_nature_cb.isChecked()
                degree_checked = hasattr(self, 'popup_degree_cb') and self.popup_degree_cb.isChecked()

                if nature_checked and degree_checked:
                    nd_text = ", which is of a nature and degree to warrant detention."
                elif nature_checked:
                    nd_text = ", which is of a nature to warrant detention."
                elif degree_checked:
                    nd_text = ", which is of a degree to warrant detention."
                else:
                    nd_text = "."

                parts.append(md_base + nd_text)

                # Nature sub-options
                if nature_checked:
                    nature_types = []
                    if hasattr(self, 'popup_relapsing_cb') and self.popup_relapsing_cb.isChecked():
                        nature_types.append("relapsing and remitting")
                    if hasattr(self, 'popup_treatment_resistant_cb') and self.popup_treatment_resistant_cb.isChecked():
                        nature_types.append("treatment resistant")
                    if hasattr(self, 'popup_chronic_cb') and self.popup_chronic_cb.isChecked():
                        nature_types.append("chronic and enduring")

                    if nature_types:
                        nature_str = ", ".join(nature_types)
                        parts.append(f"The illness is of a {nature_str} nature.")

                # Degree sub-options
                if degree_checked:
                    levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                    level = levels.get(self.popup_degree_slider.value(), "several") if hasattr(self, 'popup_degree_slider') else "several"
                    details = self.popup_degree_details.toPlainText().strip() if hasattr(self, 'popup_degree_details') else ""
                    if details:
                        parts.append(f"The degree of the illness is evidenced by {level} symptoms including {details}.")
                    else:
                        parts.append(f"The degree of the illness is evidenced by {level} symptoms.")

            elif hasattr(self, 'popup_cwd_not_met') and self.popup_cwd_not_met.isChecked():
                parts.append(md_base + ". The criteria for detention are not met.")
            else:
                parts.append(md_base + ".")

        elif hasattr(self, 'popup_md_absent') and self.popup_md_absent.isChecked():
            parts.append(f"{p['subj']} {p['does']} not suffer from a mental disorder under the Mental Health Act.")

        # 2. Necessity - Health
        if hasattr(self, 'popup_nec_yes') and self.popup_nec_yes.isChecked():
            if hasattr(self, 'popup_health_cb') and self.popup_health_cb.isChecked():
                if hasattr(self, 'popup_mental_health_cb') and self.popup_mental_health_cb.isChecked():
                    parts.append(f"Medical treatment under the Mental Health Act is necessary to prevent deterioration in {p['pos_l']} mental health.")

                    poor = hasattr(self, 'popup_poor_compliance_cb') and self.popup_poor_compliance_cb.isChecked()
                    limited = hasattr(self, 'popup_limited_insight_cb') and self.popup_limited_insight_cb.isChecked()

                    if poor and limited:
                        parts.append(f"Both historical non compliance and current limited insight makes the risk on stopping medication high without the safeguards of the Mental Health Act. This would result in a deterioration of {p['pos_l']} mental state.")
                    elif poor:
                        parts.append(f"This is based on historical non compliance and without detention I would be concerned this would result in a deterioration of {p['pos_l']} mental state.")
                    elif limited:
                        parts.append(f"I am concerned about {p['pos_l']} current limited insight into {p['pos_l']} mental health needs and how this would result in immediate non compliance with medication, hence a deterioration in {p['pos_l']} mental health.")

                if hasattr(self, 'popup_physical_health_cb') and self.popup_physical_health_cb.isChecked():
                    ph_details = self.popup_physical_health_details.toPlainText().strip() if hasattr(self, 'popup_physical_health_details') else ""
                    if ph_details:
                        parts.append(f"Medical treatment is also necessary for {p['pos_l']} physical health: {ph_details}")
                    else:
                        parts.append(f"Medical treatment is also necessary for {p['pos_l']} physical health.")

            # Safety - include checkbox details
            if hasattr(self, 'popup_safety_cb') and self.popup_safety_cb.isChecked():
                self_checked = hasattr(self, 'popup_self_cb') and self.popup_self_cb.isChecked()
                others_checked = hasattr(self, 'popup_others_cb') and self.popup_others_cb.isChecked()

                if self_checked and others_checked:
                    parts.append(f"Detention is necessary for {p['pos_l']} own safety and for the protection of others.")
                elif self_checked:
                    parts.append(f"Detention is necessary for {p['pos_l']} own safety.")
                elif others_checked:
                    parts.append("Detention is necessary for the protection of others.")

                # Self safety checkbox details - matching A3 format
                if self_checked:
                    # Determine reflexive pronoun
                    reflexive = "themselves"  # Default
                    if hasattr(self, 'popup_gender_male') and self.popup_gender_male.isChecked():
                        reflexive = "himself"
                    elif hasattr(self, 'popup_gender_female') and self.popup_gender_female.isChecked():
                        reflexive = "herself"

                    # Build risk types list: (name, is_hist, is_curr)
                    self_risk_types = [
                        ("self neglect",
                         hasattr(self, 'popup_self_neglect_hist') and self.popup_self_neglect_hist.isChecked(),
                         hasattr(self, 'popup_self_neglect_curr') and self.popup_self_neglect_curr.isChecked()),
                        (f"placing of {reflexive} in risky situations",
                         hasattr(self, 'popup_self_risky_hist') and self.popup_self_risky_hist.isChecked(),
                         hasattr(self, 'popup_self_risky_curr') and self.popup_self_risky_curr.isChecked()),
                        ("self harm",
                         hasattr(self, 'popup_self_harm_hist') and self.popup_self_harm_hist.isChecked(),
                         hasattr(self, 'popup_self_harm_curr') and self.popup_self_harm_curr.isChecked()),
                    ]

                    both_items = []
                    hist_only = []
                    curr_only = []

                    for risk_name, is_hist, is_curr in self_risk_types:
                        if is_hist and is_curr:
                            both_items.append(risk_name)
                        elif is_hist:
                            hist_only.append(risk_name)
                        elif is_curr:
                            curr_only.append(risk_name)

                    if both_items or hist_only or curr_only:
                        self_text = f"With respect to {p['pos_l']} own safety we are concerned about"
                        risk_parts = []
                        if both_items:
                            if len(both_items) == 1:
                                risk_parts.append(f"historical and current {both_items[0]}")
                            else:
                                risk_parts.append(f"historical and current {', '.join(both_items[:-1])}, and of {both_items[-1]}")
                        if hist_only:
                            if len(hist_only) == 1:
                                risk_parts.append(f"historical {hist_only[0]}")
                            else:
                                risk_parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
                        if curr_only:
                            if len(curr_only) == 1:
                                risk_parts.append(f"current {curr_only[0]}")
                            else:
                                risk_parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")

                        self_text += " " + ", and ".join(risk_parts) + "."
                        parts.append(self_text)

                    self_details = self.popup_self_details.toPlainText().strip() if hasattr(self, 'popup_self_details') else ""
                    if self_details:
                        parts.append(self_details)

                # Others safety checkbox details - matching A3 format
                if others_checked:
                    others_risk_types = [
                        ("violence to others",
                         hasattr(self, 'popup_violence_hist') and self.popup_violence_hist.isChecked(),
                         hasattr(self, 'popup_violence_curr') and self.popup_violence_curr.isChecked()),
                        ("verbal aggression",
                         hasattr(self, 'popup_verbal_hist') and self.popup_verbal_hist.isChecked(),
                         hasattr(self, 'popup_verbal_curr') and self.popup_verbal_curr.isChecked()),
                        ("sexual violence",
                         hasattr(self, 'popup_sexual_hist') and self.popup_sexual_hist.isChecked(),
                         hasattr(self, 'popup_sexual_curr') and self.popup_sexual_curr.isChecked()),
                        ("stalking",
                         hasattr(self, 'popup_stalking_hist') and self.popup_stalking_hist.isChecked(),
                         hasattr(self, 'popup_stalking_curr') and self.popup_stalking_curr.isChecked()),
                        ("arson",
                         hasattr(self, 'popup_arson_hist') and self.popup_arson_hist.isChecked(),
                         hasattr(self, 'popup_arson_curr') and self.popup_arson_curr.isChecked()),
                    ]

                    both_items = []
                    hist_only = []
                    curr_only = []

                    for risk_name, is_hist, is_curr in others_risk_types:
                        if is_hist and is_curr:
                            both_items.append(risk_name)
                        elif is_hist:
                            hist_only.append(risk_name)
                        elif is_curr:
                            curr_only.append(risk_name)

                    if both_items or hist_only or curr_only:
                        others_text = "With respect to risk to others we are concerned about the risk of"
                        risk_parts = []
                        if both_items:
                            risk_parts.append(f"historical and current {', '.join(both_items)}")
                        if hist_only:
                            risk_parts.append(f"historical {', '.join(hist_only)}")
                        if curr_only:
                            risk_parts.append(f"current {', '.join(curr_only)}")

                        others_text += " " + " and of ".join(risk_parts) + "."
                        parts.append(others_text)

                    others_details = self.popup_others_details.toPlainText().strip() if hasattr(self, 'popup_others_details') else ""
                    if others_details:
                        parts.append(others_details)

        # 3. Treatment Available & Least Restrictive - combined
        treatment_checked = hasattr(self, 'popup_treatment_cb') and self.popup_treatment_cb.isChecked()
        least_checked = hasattr(self, 'popup_least_restrictive_cb') and self.popup_least_restrictive_cb.isChecked()

        if treatment_checked and least_checked:
            parts.append(f"I can confirm appropriate medical treatment is available and this is the least restrictive option for {p['pos_l']} care.")
        elif treatment_checked:
            parts.append("I can confirm appropriate medical treatment is available.")
        elif least_checked:
            parts.append(f"This is the least restrictive option for {p['pos_l']} care.")

        # Additional comments text
        if hasattr(self, 'popup_additional'):
            additional = self.popup_additional.toPlainText().strip()
            if additional:
                parts.append(additional)

        return " ".join(parts) if parts else "No additional comments."

    def _build_popup_unfit_to_plead(self):
        """Build unfit to plead popup."""
        container, layout = self._create_popup_container("unfit_to_plead")

        label_style = "font-size: 17px; color: #374151;"
        radio_style = "QRadioButton { font-size: 17px; }"

        # First question - was patient found unfit to plead on sentencing?
        found_lbl = QLabel("Has this patient been found unfit to plead on sentencing?")
        found_lbl.setStyleSheet(label_style)
        found_lbl.setWordWrap(True)
        layout.addWidget(found_lbl)

        found_row = QHBoxLayout()
        self.popup_found_unfit_group = QButtonGroup(self)
        self.popup_found_unfit_yes = QRadioButton("Yes")
        self.popup_found_unfit_yes.setStyleSheet(radio_style)
        self.popup_found_unfit_no = QRadioButton("No")
        self.popup_found_unfit_no.setStyleSheet(radio_style)
        self.popup_found_unfit_group.addButton(self.popup_found_unfit_yes)
        self.popup_found_unfit_group.addButton(self.popup_found_unfit_no)
        found_row.addWidget(self.popup_found_unfit_yes)
        found_row.addWidget(self.popup_found_unfit_no)
        found_row.addStretch()
        layout.addLayout(found_row)

        # Container for the rest of the form (shown only if Yes)
        self.popup_unfit_details_container = QWidget()
        self.popup_unfit_details_container.setStyleSheet("background: transparent;")
        details_layout = QVBoxLayout(self.popup_unfit_details_container)
        details_layout.setContentsMargins(0, 12, 0, 0)
        details_layout.setSpacing(8)

        fit_lbl = QLabel("Is patient now fit to plead?")
        fit_lbl.setStyleSheet(label_style)
        fit_lbl.setWordWrap(True)
        details_layout.addWidget(fit_lbl)

        fit_row = QHBoxLayout()
        self.popup_fit_group = QButtonGroup(self)
        self.popup_fit_yes = QRadioButton("Yes")
        self.popup_fit_yes.setStyleSheet(radio_style)
        self.popup_fit_no = QRadioButton("No")
        self.popup_fit_no.setStyleSheet(radio_style)
        self.popup_fit_group.addButton(self.popup_fit_yes)
        self.popup_fit_group.addButton(self.popup_fit_no)
        fit_row.addWidget(self.popup_fit_yes)
        fit_row.addWidget(self.popup_fit_no)
        fit_row.addStretch()
        details_layout.addLayout(fit_row)

        details_lbl = QLabel("Details:")
        details_lbl.setStyleSheet(label_style)
        details_layout.addWidget(details_lbl)
        self.popup_unfit_details = QTextEdit()
        self.popup_unfit_details.setMaximumHeight(120)
        self.popup_unfit_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;")
        details_layout.addWidget(self.popup_unfit_details)

        self.popup_unfit_details_container.hide()
        layout.addWidget(self.popup_unfit_details_container)

        # Wire up visibility
        def on_found_unfit_changed():
            self.popup_unfit_details_container.setVisible(self.popup_found_unfit_yes.isChecked())
            self._update_preview("unfit_to_plead")

        self.popup_found_unfit_yes.toggled.connect(on_found_unfit_changed)
        self.popup_found_unfit_no.toggled.connect(on_found_unfit_changed)

        layout.addStretch()
        self._add_send_button(layout, "unfit_to_plead", self._generate_unfit)
        self._connect_preview_updates("unfit_to_plead", [
            self.popup_found_unfit_yes, self.popup_found_unfit_no,
            self.popup_fit_yes, self.popup_fit_no, self.popup_unfit_details
        ])

    def _generate_unfit(self) -> str:
        # If not found unfit to plead, return Not applicable
        if hasattr(self, 'popup_found_unfit_no') and self.popup_found_unfit_no.isChecked():
            return "Not applicable."

        # If found unfit to plead, generate the detailed response
        if hasattr(self, 'popup_found_unfit_yes') and self.popup_found_unfit_yes.isChecked():
            if self.popup_fit_yes.isChecked():
                status = "The patient is now considered fit to plead."
            elif self.popup_fit_no.isChecked():
                status = "The patient remains unfit to plead."
            else:
                status = "[Fitness to plead status]"
            details = self.popup_unfit_details.toPlainText().strip()
            return f"{status} {details}" if details else status

        return "[Select whether patient was found unfit to plead]"

    def _build_popup_signature(self):
        """Build signature popup."""
        container, layout = self._create_popup_container("signature")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        sig_lbl = QLabel("RC Signature:")
        sig_lbl.setStyleSheet(label_style)
        layout.addWidget(sig_lbl)
        self.popup_signature = QLineEdit()
        self.popup_signature.setText("Signed electronically")
        self.popup_signature.setStyleSheet(input_style)
        layout.addWidget(self.popup_signature)

        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet(label_style)
        layout.addWidget(date_lbl)
        self.popup_sig_date = QDateEdit()
        self.popup_sig_date.setCalendarPopup(True)
        self.popup_sig_date.setDate(QDate.currentDate())
        self.popup_sig_date.setDisplayFormat("dd/MM/yyyy")
        self.popup_sig_date.setStyleSheet("""
            QDateEdit {
                background: white;
                color: #1f2937;
                padding: 6px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                font-size: 17px;
            }
            QDateEdit::drop-down {
                background: #f9fafb;
                border-left: 1px solid #d1d5db;
                width: 20px;
            }
        """)
        layout.addWidget(self.popup_sig_date)

        layout.addStretch()
        self._add_send_button(layout, "signature", self._generate_signature)
        self._connect_preview_updates("signature", [self.popup_signature, self.popup_sig_date])

    def _generate_signature(self) -> str:
        sig = self.popup_signature.text()
        date = self.popup_sig_date.date().toString("dd/MM/yyyy")
        return f"Signature: {sig}\nDate: {date}"

    def _get_or_create_data_extractor(self):
        """Get or create persistent data extractor instance."""
        if self._data_extractor is None:
            try:
                from data_extractor_popup import DataExtractorPopup
                self._data_extractor = DataExtractorPopup(parent=self)
                self._data_extractor.hide()
                self._data_extractor.setWindowTitle("Data Extractor - MOJ ASR")
                self._data_extractor.setMinimumSize(800, 600)
                # Connect the data extraction signal
                if hasattr(self._data_extractor, 'data_extracted'):
                    self._data_extractor.data_extracted.connect(self._on_data_extracted)
            except ImportError:
                QMessageBox.warning(self, "Import Error", "Data extractor module not available.")
                return None
        return self._data_extractor

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            from shared_data_store import get_shared_store
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file through data extractor."""
        # Get or create persistent data extractor
        extractor = self._get_or_create_data_extractor()
        if extractor is None:
            return

        try:
            # Load the file directly into the extractor
            if hasattr(extractor, 'load_file'):
                extractor.load_file(file_path)
            elif hasattr(extractor, 'process_file'):
                extractor.process_file(file_path)
        except Exception as e:
            QMessageBox.warning(self, "Import Error", f"Failed to import file: {str(e)}")

    def _open_data_extractor(self):
        """Open data extractor for viewing/managing data (persists loaded data)."""
        extractor = self._get_or_create_data_extractor()
        if extractor is None:
            return

    def _on_data_extracted(self, data: dict):
        """Handle extracted data from data extractor."""
        # Skip if this exact data was already processed (prevents reprocessing on navigation)
        categories = data.get("categories", {})
        cat_keys = tuple(sorted(categories.keys())) if categories else ()
        cat_count = sum(len(v.get("items", [])) if isinstance(v, dict) else 0 for v in categories.values())
        content_sig = (cat_keys, cat_count)
        if self._data_processed_id == content_sig:
            print(f"[MOJ-ASR] Skipping _on_data_extracted - data already processed ({len(categories)} categories, {cat_count} items)")
            return
        self._data_processed_id = content_sig

        print(f"[MOJ-ASR] Received data with keys: {list(data.keys())}")

        # Get raw notes - first try local extractor, then fall back to SharedDataStore
        raw_notes = []
        document_type = "notes"  # Default to notes
        if hasattr(self, '_data_extractor') and self._data_extractor:
            raw_notes = getattr(self._data_extractor, 'notes', [])
            # Detect document type from the extractor's detected type
            if hasattr(self._data_extractor, '_detected_document_type'):
                document_type = self._data_extractor._detected_document_type
            # Also check if there's panel data by dtype
            elif hasattr(self._data_extractor, '_panel_data_by_dtype'):
                if self._data_extractor._panel_data_by_dtype.get("reports"):
                    document_type = "reports"
                elif self._data_extractor._panel_data_by_dtype.get("notes"):
                    document_type = "notes"

        # Fall back to SharedDataStore if no local notes
        if not raw_notes:
            try:
                from shared_data_store import get_shared_store
                shared_store = get_shared_store()
                if shared_store.has_notes():
                    raw_notes = shared_store.notes
                    print(f"[MOJ-ASR] Got {len(raw_notes)} notes from SharedDataStore (global import)")
            except Exception as e:
                print(f"[MOJ-ASR] Error getting notes from SharedDataStore: {e}")

        print(f"[MOJ-ASR] Raw notes available: {len(raw_notes)}")
        print(f"[MOJ-ASR] Document type detected: {document_type}")

        # Store raw notes at page level for section 6
        self._extracted_raw_notes = raw_notes
        self._document_type = document_type

        # Update shared data store so all sections can access these notes
        if raw_notes:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.set_notes(raw_notes, source="moj_asr_form")
            print(f"[MOJ-ASR] Updated SharedDataStore with {len(raw_notes)} notes")

        # Data extractor emits: {"categories": {"RISK": {"items": [...]}, "HPC": {"items": [...]}}}
        categories = data.get("categories", {})
        print(f"[MOJ-ASR] Categories found: {list(categories.keys())}")

        # Helper to get text from category
        def get_category_text(cat_key):
            cat_data = categories.get(cat_key, {})
            if isinstance(cat_data, dict):
                items = cat_data.get("items", [])
                texts = [item.get("text", "") for item in items if item.get("text")]
                return "\n\n".join(texts)
            elif isinstance(cat_data, str):
                return cat_data
            return ""

        # Populate section 4 behaviour import panel with HPC and RISK entries
        behaviour_entries = []

        # Check for HPC/history of presenting complaint entries (try multiple key formats)
        hpc_keys = ["HPC", "History of Presenting Complaint", "PRESENTING_COMPLAINT", "HISTORY_PRESENTING", "PRESENTING"]
        for key in hpc_keys:
            cat_data = categories.get(key, {})
            if isinstance(cat_data, dict):
                items = cat_data.get("items", [])
                for item in items:
                    if isinstance(item, dict) and item.get("text"):
                        behaviour_entries.append(item)
            elif isinstance(cat_data, list):
                for item in cat_data:
                    if isinstance(item, dict) and item.get("text"):
                        behaviour_entries.append(item)

        # Check for RISK entries (try multiple key formats)
        risk_keys = ["RISK", "Risk", "RISKS", "RISK_ASSESSMENT", "Risk Assessment"]
        for key in risk_keys:
            cat_data = categories.get(key, {})
            if isinstance(cat_data, dict):
                items = cat_data.get("items", [])
                for item in items:
                    if isinstance(item, dict) and item.get("text"):
                        behaviour_entries.append(item)
            elif isinstance(cat_data, list):
                for item in cat_data:
                    if isinstance(item, dict) and item.get("text"):
                        behaviour_entries.append(item)

        # Check for INCIDENT/FORENSIC entries
        incident_keys = ["INCIDENT", "INCIDENTS", "BEHAVIOUR", "Forensic History", "FORENSIC"]
        for key in incident_keys:
            cat_data = categories.get(key, {})
            if isinstance(cat_data, dict):
                items = cat_data.get("items", [])
                for item in items:
                    if isinstance(item, dict) and item.get("text"):
                        behaviour_entries.append(item)

        print(f"[MOJ-ASR] Found {len(behaviour_entries)} behaviour/incident entries (before date filter)")

        # Filter to last 12 months (same as section 6)
        from datetime import datetime, timedelta

        def parse_entry_date(date_val):
            if isinstance(date_val, datetime):
                return date_val
            if not date_val:
                return None
            date_str = str(date_val).strip()
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"]:
                try:
                    return datetime.strptime(date_str.split()[0] if ' ' in date_str else date_str, fmt.split()[0])
                except:
                    pass
            return None

        # Find the most recent date from raw_notes
        all_dates = []
        for n in raw_notes:
            dt = parse_entry_date(n.get("date") or n.get("datetime"))
            if dt:
                all_dates.append(dt)

        if all_dates:
            most_recent = max(all_dates)
            one_year_cutoff = most_recent - timedelta(days=365)
            print(f"[MOJ-ASR] Sections 4/5: 1-year cutoff: {one_year_cutoff.strftime('%d/%m/%Y')}")

            # Filter behaviour_entries to last 12 months
            filtered_behaviour = []
            for entry in behaviour_entries:
                entry_dt = parse_entry_date(entry.get("date") or entry.get("datetime"))
                if entry_dt and entry_dt >= one_year_cutoff:
                    entry["_sort_date"] = entry_dt
                    filtered_behaviour.append(entry)

            # Sort by date (most recent first)
            filtered_behaviour.sort(key=lambda x: x.get("_sort_date", datetime.min), reverse=True)

            # Remove temp sort key
            for e in filtered_behaviour:
                e.pop("_sort_date", None)

            print(f"[MOJ-ASR] Sections 4/5: {len(filtered_behaviour)} entries after 1-year filter")
            behaviour_entries = filtered_behaviour
        else:
            # No dates available, sort as before
            def get_sort_date(entry):
                date_str = entry.get("date", "") or entry.get("datetime", "") or ""
                return str(date_str) if date_str else ""
            behaviour_entries.sort(key=get_sort_date, reverse=True)

        # ========================================================================
        # HANDLE REPORTS vs NOTES DIFFERENTLY
        # ========================================================================
        if document_type == "reports":
            print("[MOJ-ASR] Processing as REPORT - mapping categories to sections")

            # Helper to get items from category
            def get_category_items(cat_keys):
                """Get items from multiple possible category keys."""
                items = []
                for key in cat_keys:
                    cat_data = categories.get(key, {})
                    if isinstance(cat_data, dict):
                        items.extend(cat_data.get("items", []))
                    elif isinstance(cat_data, list):
                        items.extend(cat_data)
                return items

            # Get HPC/PAST_PSYCH content for sections 4, 6, 8
            # Reports often have PAST_PSYCH or PAST PSYCH instead of HPC
            hpc_keys = ["HPC", "History of Presenting Complaint", "HISTORY_OF_PRESENTING_COMPLAINT",
                        "PRESENTING_COMPLAINT", "HISTORY_PRESENTING", "PRESENTING",
                        "PAST_PSYCH", "PAST PSYCH", "PREVIOUS_PSYCHIATRIC", "Plan"]
            hpc_items = get_category_items(hpc_keys)
            hpc_text = (get_category_text("HPC") or get_category_text("History of Presenting Complaint") or
                       get_category_text("PAST_PSYCH") or get_category_text("PAST PSYCH") or
                       get_category_text("Plan"))
            print(f"[MOJ-ASR] REPORT: Found {len(hpc_items)} HPC/PAST_PSYCH items, text length: {len(hpc_text) if hpc_text else 0}")

            # Get MSE content for section 7
            mse_keys = ["MSE", "Mental State Examination", "MENTAL_STATE", "Mental State", "MENTAL STATE",
                        "PHYSICAL HEALTH", "PHYSICAL_HEALTH"]
            mse_items = get_category_items(mse_keys)
            mse_text = (get_category_text("MSE") or get_category_text("Mental State Examination") or
                       get_category_text("MENTAL_STATE") or get_category_text("Mental State") or
                       get_category_text("PHYSICAL HEALTH"))
            print(f"[MOJ-ASR] REPORT: Found {len(mse_items)} MSE items, text length: {len(mse_text) if mse_text else 0}")

            # Get Forensic content for section 10
            forensic_keys = ["FORENSIC", "Forensic History", "FORENSIC_HISTORY", "Index Offence"]
            forensic_items = get_category_items(forensic_keys)
            forensic_text = get_category_text("FORENSIC") or get_category_text("Forensic History") or get_category_text("FORENSIC_HISTORY")
            print(f"[MOJ-ASR] REPORT: Found {len(forensic_items)} Forensic items, text length: {len(forensic_text) if forensic_text else 0}")

            # Get Risk content for section 4 (attitude/behaviour)
            risk_keys = ["Risk", "RISK", "SUMMARY", "Summary", "Risk Assessment", "RISK_ASSESSMENT"]
            risk_items = get_category_items(risk_keys)
            print(f"[MOJ-ASR] REPORT: Found {len(risk_items)} Risk/Summary items")

            # Populate Section 4 (Attitude & Behaviour) imports with HPC + Risk items combined

            section4_items = hpc_items + risk_items
            if section4_items and hasattr(self, 'populate_popup_behaviour_imports'):
                self.populate_popup_behaviour_imports(section4_items)
                print(f"[MOJ-ASR] REPORT: Section 4 imports populated with {len(section4_items)} items (HPC + Risk)")

            # Populate Section 6 (Patient's Attitude) imports with HPC/PAST_PSYCH

            if hpc_items and hasattr(self, 'populate_popup_patient_attitude_imports'):
                self.populate_popup_patient_attitude_imports(hpc_items)
                print(f"[MOJ-ASR] REPORT: Section 6 imports populated with {len(hpc_items)} items")

            # Populate Section 7 (Capacity) imports with MSE
            if mse_items and hasattr(self, 'populate_popup_capacity_imports'):
                self.populate_popup_capacity_imports(mse_items)
                print(f"[MOJ-ASR] REPORT: Section 7 imports populated with {len(mse_items)} items")

            # Populate Section 8 (Progress) imports with HPC/PAST_PSYCH
            if hpc_items and hasattr(self, 'populate_popup_progress_imports'):
                self.populate_popup_progress_imports(hpc_items)
                print(f"[MOJ-ASR] REPORT: Section 8 imports populated with {len(hpc_items)} items")

            # Populate Section 10 (How Risks Addressed) imports with Forensic
            if forensic_items and hasattr(self, 'populate_popup_risk_addressed_imports'):
                self.populate_popup_risk_addressed_imports(forensic_items)
                print(f"[MOJ-ASR] REPORT: Section 10 imports populated with {len(forensic_items)} items")

            # Extract patient details from imported report (name, DOB, gender, diagnoses, etc.)
            self._extract_patient_details_from_data()

            # ========================================================================
            # AUTO-FILL SECTION TEXT FIELDS FROM TRIBUNAL REPORT
            # ========================================================================
            prefilled_sections = []

            # Helper to get text from items
            def items_to_text(items: list) -> str:
                texts = []
                for item in items:
                    text = item.get("text", "").strip()
                    if text:
                        texts.append(text)
                return "\n\n".join(texts)

            # Section 4 (Attitude & Behaviour) - HPC + Risk content
            if hpc_items or risk_items:
                combined_items = hpc_items + risk_items
                section4_text = items_to_text(combined_items)
                if section4_text and hasattr(self, 'attitude_behaviour_text'):
                    current = self.attitude_behaviour_text.toPlainText().strip()
                    if not current:  # Only prefill if empty
                        self.attitude_behaviour_text.setPlainText(section4_text)
                        prefilled_sections.append("Section 4 (Attitude & Behaviour)")
                        print(f"[MOJ-ASR] REPORT: Auto-filled Section 4 with {len(section4_text)} chars")

            # Section 6 (Patient's Attitude) - from PAST_PSYCH or HPC
            past_psych_keys = ["PAST_PSYCH", "Past Psychiatric History", "PREVIOUS_PSYCHIATRIC", "Psychiatric History"]
            past_psych_items = get_category_items(past_psych_keys)
            attitude_items = past_psych_items if past_psych_items else hpc_items
            if attitude_items:
                section6_text = items_to_text(attitude_items)
                if section6_text and hasattr(self, 'patient_attitude_text'):
                    current = self.patient_attitude_text.toPlainText().strip()
                    if not current:
                        self.patient_attitude_text.setPlainText(section6_text)
                        prefilled_sections.append("Section 6 (Patient's Attitude)")
                        print(f"[MOJ-ASR] REPORT: Auto-filled Section 6 with {len(section6_text)} chars")

            # Section 8 (Progress) - from HPC or progress-related content
            progress_keys = ["History of Presenting Complaint", "HPC", "Progress", "PROGRESS"]
            progress_items = get_category_items(progress_keys)
            if progress_items:
                section8_text = items_to_text(progress_items)
                if section8_text and hasattr(self, 'progress_text'):
                    current = self.progress_text.toPlainText().strip()
                    if not current:
                        self.progress_text.setPlainText(section8_text)
                        prefilled_sections.append("Section 8 (Progress)")
                        print(f"[MOJ-ASR] REPORT: Auto-filled Section 8 with {len(section8_text)} chars")

            # Section 10 (How Risks Addressed) - from Forensic and Risk
            if forensic_items or risk_items:
                section10_items = forensic_items + risk_items
                section10_text = items_to_text(section10_items)
                if section10_text and hasattr(self, 'risk_factors_text'):
                    current = self.risk_factors_text.toPlainText().strip()
                    if not current:
                        self.risk_factors_text.setPlainText(section10_text)
                        prefilled_sections.append("Section 10 (Risk Factors)")
                        print(f"[MOJ-ASR] REPORT: Auto-filled Section 10 with {len(section10_text)} chars")

            section6_count = len(hpc_items)

            # Build summary message
            summary_parts = [f"Report imported. Found {len(hpc_items)} clinical items, {len(forensic_items)} forensic items."]
            if prefilled_sections:
                summary_parts.append(f"Auto-filled: {', '.join(prefilled_sections)}")
            QMessageBox.information(self, "Import Complete", "\n".join(summary_parts))

        else:
            # ========================================================================
            # HANDLE NOTES - use existing categorization and filtering
            # ========================================================================
            print("[MOJ-ASR] Processing as NOTES - using category filtering")

            print(f"[MOJ-ASR] Section 4/5: behaviour_entries count = {len(behaviour_entries)}")
            print("[MOJ-ASR] Processing Section 4 — Behaviour...")
            if behaviour_entries:
                print(f"[MOJ-ASR] Section 4/5: Populating behaviour imports with {len(behaviour_entries)} entries")
                self.populate_behaviour_import_data(behaviour_entries)
                # Also populate the popup version for the new card/popup architecture
                if hasattr(self, 'populate_popup_behaviour_imports'):
                    print(f"[MOJ-ASR] Section 4 popup: Calling populate_popup_behaviour_imports")
                    self.populate_popup_behaviour_imports(behaviour_entries)
                # Also populate section 5 imported data (same data relevant for addressing issues)
                self.populate_section5_import_data(behaviour_entries)
            else:
                print("[MOJ-ASR] Section 4/5: No behaviour_entries found - using raw_notes instead")
                # Fall back to using raw_notes if no categorized behaviour entries
                if hasattr(self, 'populate_popup_behaviour_imports') and raw_notes:
                    print(f"[MOJ-ASR] Section 4 popup: Calling populate_popup_behaviour_imports with {len(raw_notes)} raw notes")
                    self.populate_popup_behaviour_imports(raw_notes)

            # Populate section 6 with mental state, compliance, attendance, admissions (1 year filter)
            # This also builds the timeline which Section 4 needs for admission tagging
            print("[MOJ-ASR] Processing Section 6 — Timeline & Admissions...")
            section6_count = self._populate_section6_filtered_data()

            # Re-populate Section 4 popup now that timeline is available for admission tagging
            if hasattr(self, '_timeline_episodes') and self._timeline_episodes:
                print(f"[MOJ-ASR] Re-populating Section 4 popup with timeline data ({len(self._timeline_episodes)} episodes)")
                entries_for_s4 = behaviour_entries if behaviour_entries else raw_notes
                if hasattr(self, 'populate_popup_behaviour_imports') and entries_for_s4:
                    self.populate_popup_behaviour_imports(entries_for_s4)

            # Populate Section 6 popup (Patient's Attitude) with offending behaviour related notes
            print("[MOJ-ASR] Processing Section 6 — Patient Attitude...")
            entries_for_s6 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_patient_attitude_imports') and entries_for_s6:
                print(f"[MOJ-ASR] Section 6 popup: Calling populate_popup_patient_attitude_imports with {len(entries_for_s6)} entries")
                self.populate_popup_patient_attitude_imports(entries_for_s6)

            # Populate Section 7 popup (Capacity) with capacity-related notes
            print("[MOJ-ASR] Processing Section 7 — Capacity...")
            entries_for_s7 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_capacity_imports') and entries_for_s7:
                print(f"[MOJ-ASR] Section 7 popup: Calling populate_popup_capacity_imports with {len(entries_for_s7)} entries")
                self.populate_popup_capacity_imports(entries_for_s7)

            # Populate Section 10 popup (How Risks Addressed) with risk factor notes
            print("[MOJ-ASR] Processing Section 10 — Risk Factors...")
            entries_for_s10 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_risk_addressed_imports') and entries_for_s10:
                print(f"[MOJ-ASR] Section 10 popup: Calling populate_popup_risk_addressed_imports with {len(entries_for_s10)} entries")
                self.populate_popup_risk_addressed_imports(entries_for_s10)

            # Populate Section 11 popup (Abscond/Escape) with AWOL-related notes
            print("[MOJ-ASR] Processing Section 11 — Abscond/Escape...")
            entries_for_s11 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_abscond_imports') and entries_for_s11:
                print(f"[MOJ-ASR] Section 11 popup: Calling populate_popup_abscond_imports with {len(entries_for_s11)} entries")
                self.populate_popup_abscond_imports(entries_for_s11)

            # Populate Section 12 popup (MAPPA) with MAPPA-related notes
            print("[MOJ-ASR] Processing Section 12 — MAPPA...")
            entries_for_s12 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_mappa_imports') and entries_for_s12:
                print(f"[MOJ-ASR] Section 12 popup: Calling populate_popup_mappa_imports with {len(entries_for_s12)} entries")
                self.populate_popup_mappa_imports(entries_for_s12)

            # Populate Section 14 popup (Leave Report) with leave-related notes
            print("[MOJ-ASR] Processing Section 14 — Leave Report...")
            entries_for_s14 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_leave_report_imports') and entries_for_s14:
                print(f"[MOJ-ASR] Section 14 popup: Calling populate_popup_leave_report_imports with {len(entries_for_s14)} entries")
                self.populate_popup_leave_report_imports(entries_for_s14)

            # Populate Section 8 popup (Progress) with mental state/progress-related notes
            print("[MOJ-ASR] Processing Section 8 — Progress...")
            entries_for_s8 = behaviour_entries if behaviour_entries else raw_notes
            if hasattr(self, 'populate_popup_progress_imports') and entries_for_s8:
                print(f"[MOJ-ASR] Section 8 popup: Calling populate_popup_progress_imports with {len(entries_for_s8)} entries")
                self.populate_popup_progress_imports(entries_for_s8)

            # Extract patient details from imported data (name, DOB, gender, etc.)
            print("[MOJ-ASR] Extracting patient details...")
            self._extract_patient_details_from_data()

            QMessageBox.information(self, "Import Complete", f"Data imported. {len(behaviour_entries)} entries for sections 4/5, {section6_count} entries for section 6 (last 12 months).")

    def _extract_patient_details_from_data(self):
        """Extract patient details (name, DOB, gender, hospital, MHA) from imported documents."""
        import re
        from datetime import datetime

        # Get document text from multiple sources
        document_text = ""
        notes = []

        # Source 1: Data extractor notes
        if hasattr(self, '_data_extractor') and self._data_extractor:
            extractor_notes = getattr(self._data_extractor, 'notes', [])
            if not extractor_notes and hasattr(self._data_extractor, '_raw_notes'):
                extractor_notes = self._data_extractor._raw_notes
            notes.extend(extractor_notes or [])

        # Source 2: Page-level extracted raw notes
        if hasattr(self, '_extracted_raw_notes') and self._extracted_raw_notes:
            notes.extend(self._extracted_raw_notes)

        # Combine text from all notes
        for note in notes:
            text = note.get("text", "") or note.get("body", "") or note.get("content", "")
            if text:
                document_text += text + "\n"

        if not document_text.strip():
            print("[MOJ-ASR] No document text for patient details extraction")
            return

        print(f"[MOJ-ASR] Extracting patient details from {len(document_text)} chars ({len(notes)} notes)")
        # Debug: show first 500 chars
        print(f"[MOJ-ASR] Document preview: {document_text[:500]}...")

        # ================================================================
        # EXTRACT PATIENT NAME
        # ================================================================
        # Medical qualifications to exclude from name matching
        MEDICAL_QUALIFICATIONS = {
            'mrcp', 'mrcpsych', 'mbbs', 'frcp', 'frcpsych', 'md', 'phd',
            'bsc', 'msc', 'ba', 'ma', 'dpm', 'mphil', 'dclinpsy', 'ccst',
            'gmc', 'nmc', 'hcpc', 'consultant', 'registrar', 'sho'
        }

        name_patterns = [
            # Most specific: "Full Name:" pattern (common in tribunal reports)
            (r'full\s+name\s*[:\s]+(?:Mr|Mrs|Ms|Miss|Dr)?\.?\s*([A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\']+){1,3})', 'full_name'),
            # Patient details section with name
            (r'patient\s+details.*?(?:full\s+)?name\s*[:\s]+(?:Mr|Mrs|Ms|Miss|Dr)?\.?\s*([A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\']+){1,3})', 'patient_details_name'),
            # Standard patterns
            (r'^([A-Z][A-Z\-\']+),\s*([A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\']+)?)', 'surname_first_header'),
            (r'patient\s+name\s*[:]?\s*([A-Z][a-zA-Z\-\']+(?:[ ]+[A-Z][a-zA-Z\-\']+){0,3})(?=\n|$)', 'patient_name_label'),
            (r'name\s+of\s+patient\s*[:]?\s*\n\s*([A-Z][a-zA-Z\-\']+(?:[ ]+[A-Z][a-zA-Z\-\']+){0,3})(?=\n|$)', 'name_of_patient'),
            (r'(?:patient\s+)?name\s*[:]\s*([A-Z][a-zA-Z\-\']+(?:[ ]+[A-Z][a-zA-Z\-\']+){0,3})(?=\n|$)', 'name_colon'),
            (r're\s*[:]\s*([A-Z][a-zA-Z\-\']+(?:[ ]+[A-Z][a-zA-Z\-\']+){0,3})(?=\n|$)', 're_colon'),
            (r'\b([A-Z][A-Z\-\']+),\s+([A-Z][a-zA-Z\-\']+(?:\s+[A-Z][a-zA-Z\-\']+)?)\s*(?:\(|$|\n)', 'surname_comma_first'),
            (r'\b((?:Mr|Mrs|Ms|Miss|Dr)\.?\s+[A-Z][a-zA-Z\-\']+(?:[ ]+[A-Z][a-zA-Z\-\']+){0,2})\b', 'title_name'),
        ]

        extracted_name = None
        for pattern, pattern_name in name_patterns:
            flags = re.IGNORECASE | re.DOTALL if pattern_name in ('full_name', 'patient_details_name') else (re.IGNORECASE if 'header' not in pattern_name else 0)
            match = re.search(pattern, document_text, flags)
            if match:
                if pattern_name in ('surname_first_header', 'surname_comma_first'):
                    surname = match.group(1).strip()
                    firstname = match.group(2).strip()
                    # Skip if this looks like medical qualifications
                    if surname.lower() in MEDICAL_QUALIFICATIONS or firstname.lower() in MEDICAL_QUALIFICATIONS:
                        print(f"[MOJ-ASR] Skipping medical qualification match: {surname}, {firstname}")
                        continue
                    if surname.isupper():
                        surname = surname.title()
                    name = f"{firstname} {surname}"
                else:
                    name = match.group(1).strip()
                # Clean up the name
                name = re.sub(r'\s*\(.*?\)\s*', ' ', name)
                name = re.sub(r'[,;:\.\s]+$', '', name).strip()
                name = re.sub(r'\s+', ' ', name)  # Normalize whitespace
                # Validate: not too short, not too long, not a medical qualification
                name_words = name.lower().split()
                if any(word in MEDICAL_QUALIFICATIONS for word in name_words):
                    print(f"[MOJ-ASR] Skipping name with medical qualification: {name}")
                    continue
                if 2 <= len(name) <= 60 and len(name_words) >= 2:
                    extracted_name = name
                    print(f"[MOJ-ASR] Extracted patient name: '{extracted_name}' (pattern: {pattern_name})")
                    break

        # ================================================================
        # EXTRACT DATE OF BIRTH
        # ================================================================
        dob_patterns = [
            # "DATE OF BIRTH" followed by date on same line or next (with or without colon)
            r'date\s+of\s+birth\s*[:]?\s*\n?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'date\s+of\s+birth\s*[:]?\s*\n?\s*(\d{1,2}\s+[A-Za-z]+\s+\d{4})',
            # "DOB" followed by date
            r'd\.?o\.?b\.?\s*[:]\s*\n?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'd\.?o\.?b\.?\s*[:]\s*\n?\s*(\d{1,2}\s+[A-Za-z]+\s+\d{4})',
            # Date followed by "(X years)" pattern - common in headers like "4 Jun 1985 (40 years)"
            r'(\d{1,2}\s+[A-Za-z]+\s+\d{4})\s*\(\d+\s*years?\)',
            # Standalone date patterns near "birth" or "born"
            r'born\s*[:]?\s*\n?\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
        ]

        extracted_dob = None
        for pattern in dob_patterns:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                print(f"[MOJ-ASR] DOB pattern matched: {pattern[:50]}... -> '{match.group(1)}'")
                date_str = match.group(1).strip()
                for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y", "%d %B %Y", "%d %b %Y"]:
                    try:
                        parsed_date = datetime.strptime(date_str, fmt)
                        if parsed_date.year > 2050:
                            parsed_date = parsed_date.replace(year=parsed_date.year - 100)
                        extracted_dob = parsed_date
                        print(f"[MOJ-ASR] Extracted DOB: {extracted_dob.strftime('%d/%m/%Y')}")
                        break
                    except ValueError:
                        continue
                if extracted_dob:
                    break

        # ================================================================
        # EXTRACT GENDER
        # ================================================================
        extracted_gender = None
        if re.search(r'(?<![a-zA-Z])Female(?![a-z])', document_text):
            extracted_gender = "Female"
        elif re.search(r'(?<![a-zA-Z])Male(?![a-z])', document_text):
            if not re.search(r'(?<![a-zA-Z])Female(?![a-z])', document_text):
                extracted_gender = "Male"
        if not extracted_gender:
            if re.search(r'\(\s*(?:Miss|Mrs|Ms)\s*\)', document_text, re.IGNORECASE):
                extracted_gender = "Female"
            elif re.search(r'\(\s*Mr\s*\)', document_text, re.IGNORECASE):
                extracted_gender = "Male"
        if not extracted_gender:
            if re.search(r'\bMr\.?\s+[A-Z]', document_text):
                extracted_gender = "Male"
            elif re.search(r'\b(?:Mrs|Ms|Miss)\.?\s+[A-Z]', document_text):
                extracted_gender = "Female"
        if not extracted_gender:
            female_pronouns = len(re.findall(r'\b(?:she|her|herself)\b', document_text.lower()))
            male_pronouns = len(re.findall(r'\b(?:he|him|himself)\b', document_text.lower()))
            if female_pronouns > male_pronouns * 2 and female_pronouns >= 5:
                extracted_gender = "Female"
            elif male_pronouns > female_pronouns * 2 and male_pronouns >= 5:
                extracted_gender = "Male"
        if extracted_gender:
            print(f"[MOJ-ASR] Extracted gender: {extracted_gender}")

        # ================================================================
        # EXTRACT HOSPITAL
        # ================================================================
        hospital_patterns = [
            # "Usual Place of Residence:" - can be on same line or next line
            r'usual\s+place\s+of\s+residence\s*[:\s]+([A-Za-z][^\n]+?)(?=\n|$)',
            # "Hospital:" pattern
            r'(?:detaining\s+)?hospital\s*[:\s]+([A-Za-z][A-Za-z\s\-\'\.]+?)(?=\n|$)',
            # "Admitted to:" pattern
            r'admitted\s+to\s*[:\s]+([A-Za-z][A-Za-z\s\-\'\.]+?)(?=\n|$)',
            # Hospital name followed by postcode
            r'([A-Za-z][A-Za-z\s\-\']+(?:Hospital|Unit|Centre|Center|Ward)[A-Za-z\s\-\']*)\s+[A-Z]{1,2}\d{1,2}\s*\d[A-Z]{2}',
        ]
        extracted_hospital = None
        for pattern in hospital_patterns:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                hospital = match.group(1).strip()
                # Clean up trailing punctuation
                hospital = re.sub(r'[,;:\.\s]+$', '', hospital).strip()
                if hospital and len(hospital) >= 3:
                    extracted_hospital = hospital
                    print(f"[MOJ-ASR] Extracted hospital: '{extracted_hospital}'")
                    break

        # ================================================================
        # EXTRACT MHA SECTION
        # ================================================================
        mha_patterns = [
            r'mental\s+health\s+act\s+(?:status|section)\s*[:]?\s*\n\s*(S?\d+[\/\-]?\d*[A-Za-z]?)',
            r'(?:mha\s+)?section\s*[:]\s*\n?\s*(S?\d+[\/\-]?\d*[A-Za-z]?)',
            r'\b(S?37[\/\-]41|S?47[\/\-]49|S?48[\/\-]49|S?45A)\b',
        ]
        extracted_mha = None
        for pattern in mha_patterns:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                mha = match.group(1).strip()
                mha = re.sub(r'^S', '', mha).replace('-', '/')
                extracted_mha = mha
                print(f"[MOJ-ASR] Extracted MHA: '{extracted_mha}'")
                break

        # ================================================================
        # EXTRACT MHA SECTION START DATE (from text near section mention)
        # ================================================================
        extracted_mha_date = None
        if extracted_mha:
            # Build search patterns that capture date near section mention
            # e.g., "detained under s.37/41 on 29 Nov 2007"
            mha_search = extracted_mha.replace('/', '[/\\-]')  # Match 37/41 or 37-41

            # Patterns to find date AFTER section mention
            date_after_patterns = [
                # "s.37/41 on 29 Nov 2007" or "section 37/41 on 29/11/2007"
                rf'(?:section|s\.?)\s*{mha_search}\s+(?:on|from|since|dated?)\s+(\d{{1,2}}\s+[A-Za-z]+\s+\d{{4}})',
                rf'(?:section|s\.?)\s*{mha_search}\s+(?:on|from|since|dated?)\s+(\d{{1,2}}[/\-\.]\d{{1,2}}[/\-\.]\d{{2,4}})',
                # "detained under s.37/41 on 29 Nov 2007"
                rf'(?:detained|placed|sectioned)\s+(?:under\s+)?(?:section|s\.?)\s*{mha_search}\s+(?:on|from|since)?\s*(\d{{1,2}}\s+[A-Za-z]+\s+\d{{4}})',
                rf'(?:detained|placed|sectioned)\s+(?:under\s+)?(?:section|s\.?)\s*{mha_search}\s+(?:on|from|since)?\s*(\d{{1,2}}[/\-\.]\d{{1,2}}[/\-\.]\d{{2,4}})',
                # "37/41 since 29 Nov 2007"
                rf'\b{mha_search}\s+(?:on|from|since|dated?)\s+(\d{{1,2}}\s+[A-Za-z]+\s+\d{{4}})',
                rf'\b{mha_search}\s+(?:on|from|since|dated?)\s+(\d{{1,2}}[/\-\.]\d{{1,2}}[/\-\.]\d{{2,4}})',
            ]

            # Patterns to find date BEFORE section mention
            date_before_patterns = [
                # "on 29 Nov 2007 detained under s.37/41"
                rf'(?:on|from|since)\s+(\d{{1,2}}\s+[A-Za-z]+\s+\d{{4}})\s+(?:detained|placed|sectioned)\s+(?:under\s+)?(?:section|s\.?)\s*{mha_search}',
                rf'(?:on|from|since)\s+(\d{{1,2}}[/\-\.]\d{{1,2}}[/\-\.]\d{{2,4}})\s+(?:detained|placed|sectioned)\s+(?:under\s+)?(?:section|s\.?)\s*{mha_search}',
            ]

            all_patterns = date_after_patterns + date_before_patterns

            for pattern in all_patterns:
                match = re.search(pattern, document_text, re.IGNORECASE)
                if match:
                    date_str = match.group(1).strip()
                    print(f"[MOJ-ASR] Found MHA date pattern: '{date_str}'")

                    # Parse the date
                    for fmt in ["%d %B %Y", "%d %b %Y", "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y",
                                "%d/%m/%y", "%d-%m-%y"]:
                        try:
                            parsed_date = datetime.strptime(date_str, fmt)
                            # Handle 2-digit years
                            if parsed_date.year > 2050:
                                parsed_date = parsed_date.replace(year=parsed_date.year - 100)
                            extracted_mha_date = parsed_date
                            print(f"[MOJ-ASR] Extracted MHA Section Date: {extracted_mha_date.strftime('%d/%m/%Y')}")
                            break
                        except:
                            pass

                    if extracted_mha_date:
                        break

        # ================================================================
        # EXTRACT NHS NUMBER
        # ================================================================
        nhs_patterns = [
            # "NHS Number:" on one line, number on next
            r'nhs\s+(?:number|no\.?)\s*[:]?\s*\n\s*(\d[\d\s]{8,12}\d)',
            # "NHS Number:" followed by number on same line
            r'nhs\s+(?:number|no\.?)\s*[:]\s*(\d[\d\s]{8,12}\d)',
            # "Hospital Number:" patterns
            r'hospital\s+(?:number|no\.?)\s*[:]?\s*\n?\s*(\d[\d\s]{6,15})',
        ]
        extracted_nhs = None
        for pattern in nhs_patterns:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                nhs = match.group(1).strip()
                extracted_nhs = nhs
                print(f"[MOJ-ASR] Extracted NHS Number: '{extracted_nhs}'")
                break

        # ================================================================
        # EXTRACT DIAGNOSES (match to ICD-10) - copied from Leave Form
        # ================================================================
        extracted_diagnoses = []
        matched_categories = set()  # Track matched categories to prevent duplicates

        # Import ICD10_DICT
        try:
            from icd10_dict import load_icd10_dict
            ICD10_DICT = load_icd10_dict()
        except (ImportError, FileNotFoundError):
            ICD10_DICT = {}

        # Clean the document text for diagnosis extraction:
        # 1. Remove question headings (e.g., "Does the patient have a learning disability?")
        # 2. Remove form prompts that might cause false matches
        doc_for_diagnosis = document_text

        # Remove common question patterns that aren't actual diagnoses
        question_patterns_to_remove = [
            r'does\s+the\s+patient\s+have\s+a\s+learning\s+disabilit[y]?\s*\?',
            r'is\s+the\s+patient\s+(?:now\s+)?suffering\s+from\s+(?:a\s+)?mental\s+disorder\s*\?',
            r'does\s+the\s+patient\s+have\s+(?:a\s+)?(?:diagnosis\s+of\s+)?',
            r'\d+\.\s*(?:does|is|has|are)\s+the\s+patient\s+[^\n]+\?',  # Numbered questions
        ]
        for pattern in question_patterns_to_remove:
            doc_for_diagnosis = re.sub(pattern, '', doc_for_diagnosis, flags=re.IGNORECASE)

        doc_lower = doc_for_diagnosis.lower()

        # Handle "revised to" or "changed to" - prefer the current diagnosis
        # e.g., "Paranoid Schizophrenia that was revised to Schizoaffective Disorder"
        revised_patterns = [
            r'(\w+(?:\s+\w+)*)\s+(?:that\s+was\s+)?(?:revised|changed|updated|modified)\s+to\s+(\w+(?:\s+\w+)*)',
        ]
        superseded_diagnoses = set()
        for pattern in revised_patterns:
            for match in re.finditer(pattern, doc_lower):
                old_diag = match.group(1).strip()
                # Mark the old diagnosis as superseded
                superseded_diagnoses.add(old_diag)
                print(f"[MOJ-ASR] Diagnosis '{old_diag}' superseded (revised to something else)")

        def find_icd10_entry(search_term, course_preference="continuous"):
            """Find ICD-10 entry, preferring specified course type."""
            search_lower = search_term.lower()
            best_match = None
            preferred_match = None

            for diag_name, meta in ICD10_DICT.items():
                diag_lower = diag_name.lower()
                if search_lower in diag_lower:
                    icd_code = meta.get("icd10", "") if isinstance(meta, dict) else ""
                    display = f"{diag_name} ({icd_code})" if icd_code else diag_name

                    # Check for preferred course type
                    if course_preference == "continuous" and "continuous" in diag_lower:
                        preferred_match = display
                        break
                    elif course_preference == "remission" and "remission" in diag_lower:
                        preferred_match = display
                        break
                    elif best_match is None:
                        best_match = display

            return preferred_match or best_match

        # Check for late autism diagnosis indicators
        late_autism_indicators = [
            "late diagnos", "recently diagnos", "adult diagnos", "diagnosed as an adult",
            "diagnosed in adult", "diagnosed late", "diagnosis was made in", "diagnosed at age",
            "diagnosed aged", "diagnosed when he was", "diagnosed when she was"
        ]
        is_late_autism = any(indicator in doc_lower for indicator in late_autism_indicators)

        # Also check for autism diagnosis date - if diagnosed as adult (age 18+), it's atypical
        if not is_late_autism and extracted_dob:
            # Look for "diagnosed with autism in [month] [year]" or "autism... diagnosed in [year]"
            autism_diag_match = re.search(
                r'(?:diagnosed\s+with\s+autism|autism.*?diagnos\w*)\s+(?:in\s+)?'
                r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+)?(\d{4})',
                doc_lower
            )
            if autism_diag_match:
                diag_year = int(autism_diag_match.group(1))
                birth_year = extracted_dob.year
                age_at_diagnosis = diag_year - birth_year
                print(f"[MOJ-ASR] Autism diagnosed in {diag_year}, patient born {birth_year}, age at diagnosis: ~{age_at_diagnosis}")
                if age_at_diagnosis >= 18:
                    is_late_autism = True
                    print(f"[MOJ-ASR] Late autism diagnosis detected (age {age_at_diagnosis} >= 18)")

        # Check for schizophrenia course indicators
        def get_schizophrenia_course():
            if "complete remission" in doc_lower or "in remission" in doc_lower:
                return "remission"
            elif "episodic" in doc_lower:
                return "episodic"
            else:
                return "continuous"  # Default to continuous

        schizo_course = get_schizophrenia_course()

        # Diagnosis patterns with category tracking
        DIAGNOSIS_PATTERNS = [
            # Schizophrenia variants (category: schizophrenia)
            (r'paranoid\s+schizophrenia', 'Paranoid schizophrenia', 'schizophrenia'),
            (r'catatonic\s+schizophrenia', 'Catatonic schizophrenia', 'schizophrenia'),
            (r'hebephrenic\s+schizophrenia', 'Hebephrenic schizophrenia', 'schizophrenia'),
            (r'residual\s+schizophrenia', 'Residual schizophrenia', 'schizophrenia'),
            (r'simple\s+schizophrenia', 'Simple schizophrenia', 'schizophrenia'),
            (r'undifferentiated\s+schizophrenia', 'Undifferentiated schizophrenia', 'schizophrenia'),
            (r'schizoaffective', 'Schizoaffective disorder', 'schizoaffective'),
            (r'schizophrenia', 'Schizophrenia, unspecified', 'schizophrenia'),
            # Autism (category: autism) - check late diagnosis
            (r'autism\s+spectrum\s+disorder', None, 'autism'),  # Will be resolved dynamically
            (r'autistic\s+spectrum', None, 'autism'),
            (r'asd\b', None, 'autism'),
            (r'asperger', 'Asperger', 'autism'),
            (r'atypical\s+autism', 'Atypical autism', 'autism'),
            # Mood disorders
            (r'bipolar\s+affective\s+disorder', 'Bipolar affective disorder', 'bipolar'),
            (r'bipolar\s+disorder', 'Bipolar affective disorder', 'bipolar'),
            (r'manic\s+depressi', 'Bipolar affective disorder', 'bipolar'),
            (r'recurrent\s+depressi', 'Recurrent depressive disorder', 'depression'),
            (r'major\s+depressi', 'Depressive episode', 'depression'),
            (r'depressi(?:ve|on)', 'Depressive episode', 'depression'),
            # Personality disorders
            (r'emotionally\s+unstable\s+personality', 'Emotionally unstable personality disorder', 'personality'),
            (r'borderline\s+personality', 'Emotionally unstable personality disorder', 'personality'),
            (r'antisocial\s+personality', 'Dissocial personality disorder', 'personality'),
            (r'dissocial\s+personality', 'Dissocial personality disorder', 'personality'),
            (r'narcissistic\s+personality', 'Other specific personality disorders', 'personality'),
            (r'paranoid\s+personality', 'Paranoid personality disorder', 'personality'),
            (r'personality\s+disorder', 'Personality disorder', 'personality'),
            # Anxiety
            (r'generalised\s+anxiety', 'Generalized anxiety disorder', 'anxiety'),
            (r'generalized\s+anxiety', 'Generalized anxiety disorder', 'anxiety'),
            (r'ptsd', 'Post-traumatic stress disorder', 'ptsd'),
            (r'post[- ]?traumatic\s+stress', 'Post-traumatic stress disorder', 'ptsd'),
            # Psychosis
            (r'acute.*psycho(?:tic|sis)', 'Acute and transient psychotic disorders', 'psychosis'),
            # Learning disability - require diagnosis context, not just mention
            (r'(?:diagnosis\s+of|diagnosed\s+with|has\s+(?:a\s+)?|suffers?\s+from)\s+(?:mild\s+|moderate\s+|severe\s+)?learning\s+disabilit', 'Mental retardation', 'learning'),
            (r'(?:diagnosis\s+of|diagnosed\s+with|has\s+(?:a\s+)?|suffers?\s+from)\s+(?:mild\s+|moderate\s+|severe\s+)?intellectual\s+disabilit', 'Mental retardation', 'learning'),
            (r'learning\s+disability\s*(?:\(|F7)', 'Mental retardation', 'learning'),  # With ICD code
            # Substance
            (r'alcohol\s+dependence', 'Alcohol dependence syndrome', 'alcohol'),
            (r'drug\s+dependence', 'Drug dependence', 'drugs'),
            (r'opioid\s+dependence', 'Opioid dependence', 'drugs'),
        ]

        if ICD10_DICT:
            for pattern, search_term, category in DIAGNOSIS_PATTERNS:
                # Skip if we've already matched this category
                if category in matched_categories:
                    continue

                match = re.search(pattern, doc_lower)
                if match:
                    # Check if this diagnosis term appears in a superseded context
                    matched_text = match.group(0).strip()
                    is_superseded = False
                    for superseded in superseded_diagnoses:
                        if superseded in matched_text or matched_text in superseded:
                            print(f"[MOJ-ASR] Skipping superseded diagnosis: '{matched_text}'")
                            is_superseded = True
                            break
                    if is_superseded:
                        continue

                    # Handle autism specially - check for late diagnosis
                    if category == 'autism' and search_term is None:
                        if is_late_autism:
                            search_term = 'Atypical autism'
                            print(f"[MOJ-ASR] Late autism diagnosis detected -> Atypical autism")
                        else:
                            search_term = 'Childhood autism'

                    # Handle schizophrenia - use appropriate course
                    if category == 'schizophrenia':
                        # Check if this specific type was superseded
                        if any(search_term.lower() in s for s in superseded_diagnoses):
                            print(f"[MOJ-ASR] Skipping superseded schizophrenia type: '{search_term}'")
                            continue
                        display = find_icd10_entry(search_term, schizo_course)
                        print(f"[MOJ-ASR] Schizophrenia with course '{schizo_course}' -> {display}")
                    else:
                        display = find_icd10_entry(search_term)

                    if display and display not in extracted_diagnoses:
                        extracted_diagnoses.append(display)
                        matched_categories.add(category)
                        print(f"[MOJ-ASR] Matched diagnosis: '{pattern}' -> '{display}'")

                    # Limit to 3 diagnoses
                    if len(extracted_diagnoses) >= 3:
                        break

        print(f"[MOJ-ASR] Extracted {len(extracted_diagnoses)} diagnoses")

        # ================================================================
        # EXTRACT COMPLIANCE/UNDERSTANDING FOR SECTION 6 (like Leave 4g)
        # ================================================================
        # Treatment-specific compliance keywords
        UNDERSTANDING_KEYWORDS = {
            "good": ["good understanding", "understands well", "good insight", "clear understanding",
                    "demonstrates understanding", "aware of", "recognises"],
            "fair": ["fair understanding", "some understanding", "partial understanding",
                    "limited understanding", "developing understanding"],
            "poor": ["poor understanding", "no understanding", "lacks understanding",
                    "does not understand", "limited insight", "poor insight"],
        }

        COMPLIANCE_KEYWORDS = {
            "full": ["fully compliant", "full compliance", "excellent compliance", "good compliance",
                    "complies with", "compliant with medication", "takes medication as prescribed",
                    "engaging well", "engages well", "good engagement"],
            "reasonable": ["reasonable compliance", "reasonably compliant", "generally compliant",
                         "mostly compliant", "adequate compliance", "fair compliance"],
            "partial": ["partial compliance", "partially compliant", "variable compliance",
                       "inconsistent compliance", "sometimes non-compliant", "irregular"],
            "nil": ["non-compliant", "not compliant", "refuses", "refusing", "poor compliance",
                   "non-compliance", "disengaged", "not engaging"],
        }

        # Treatment-specific keywords to detect which treatment is being discussed
        TREATMENT_CONTEXT = {
            "medical": ["medical", "medication", "depot", "clozapine", "antipsychotic", "psychiatr"],
            "nursing": ["nursing", "nurse", "ward staff", "observations", "level of obs"],
            "psychology": ["psychology", "psycholog", "psychological", "therapy sessions", "1-1 work"],
            "ot": ["occupational therapy", "ot ", " ot,", "groups", "activities", "horticulture", "cooking"],
            "social_work": ["social work", "social worker", "discharge", "accommodation", "family work"],
        }

        extracted_treatment_compliance = {}

        for treatment, context_keywords in TREATMENT_CONTEXT.items():
            # Find text segments mentioning this treatment
            for context_kw in context_keywords:
                if context_kw in doc_lower:
                    # Look for compliance/understanding near this context
                    # Check understanding
                    for level, keywords in UNDERSTANDING_KEYWORDS.items():
                        for kw in keywords:
                            if kw in doc_lower:
                                if treatment not in extracted_treatment_compliance:
                                    extracted_treatment_compliance[treatment] = {}
                                if "understanding" not in extracted_treatment_compliance[treatment]:
                                    extracted_treatment_compliance[treatment]["understanding"] = level
                                    print(f"[MOJ-ASR] Section 6: {treatment} understanding = {level}")
                                break
                        if treatment in extracted_treatment_compliance and "understanding" in extracted_treatment_compliance[treatment]:
                            break

                    # Check compliance
                    for level, keywords in COMPLIANCE_KEYWORDS.items():
                        for kw in keywords:
                            if kw in doc_lower:
                                if treatment not in extracted_treatment_compliance:
                                    extracted_treatment_compliance[treatment] = {}
                                if "compliance" not in extracted_treatment_compliance[treatment]:
                                    extracted_treatment_compliance[treatment]["compliance"] = level
                                    print(f"[MOJ-ASR] Section 6: {treatment} compliance = {level}")
                                break
                        if treatment in extracted_treatment_compliance and "compliance" in extracted_treatment_compliance[treatment]:
                            break
                    break

        # Apply defaults if nothing specific found - look for general compliance
        general_compliance = None
        general_understanding = None

        for level, keywords in COMPLIANCE_KEYWORDS.items():
            for kw in keywords:
                if kw in doc_lower:
                    general_compliance = level
                    break
            if general_compliance:
                break

        for level, keywords in UNDERSTANDING_KEYWORDS.items():
            for kw in keywords:
                if kw in doc_lower:
                    general_understanding = level
                    break
            if general_understanding:
                break

        # Default to "good" understanding and "full" compliance if nothing found
        if not general_compliance:
            general_compliance = "full"
        if not general_understanding:
            general_understanding = "good"

        # Apply general to all treatments that don't have specific values
        for treatment in TREATMENT_CONTEXT.keys():
            if treatment not in extracted_treatment_compliance:
                extracted_treatment_compliance[treatment] = {}
            if "understanding" not in extracted_treatment_compliance[treatment]:
                extracted_treatment_compliance[treatment]["understanding"] = general_understanding
            if "compliance" not in extracted_treatment_compliance[treatment]:
                extracted_treatment_compliance[treatment]["compliance"] = general_compliance

        print(f"[MOJ-ASR] Section 6: Extracted compliance for {len(extracted_treatment_compliance)} treatments")

        # ================================================================
        # EXTRACT AGE (from pattern or calculate from DOB)
        # ================================================================
        extracted_age = None
        age_patterns = [
            r"(?:AGE)[:\s]*(\d{1,3})\s*(?:years?|yrs?|y\.?o\.?)?\b",
            r"\b(\d{1,3})\s*(?:year|yr)\s*old\b",
            r"\b(\d{1,3})\s*y\.?o\.?\b",
            r"\baged?\s*(\d{1,3})\b",
            r'\(\s*(\d{1,3})\s*years?\s*\)',  # "(40 years)" pattern common in headers
        ]
        for pattern in age_patterns:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                age_val = int(match.group(1))
                if 0 < age_val < 120:
                    extracted_age = age_val
                    print(f"[MOJ-ASR] Extracted age from pattern: {extracted_age}")
                    break

        # Calculate age from DOB if not found explicitly
        if not extracted_age and extracted_dob:
            today = datetime.today()
            age = today.year - extracted_dob.year - ((today.month, today.day) < (extracted_dob.month, extracted_dob.day))
            if 0 < age < 120:
                extracted_age = age
                print(f"[MOJ-ASR] Calculated age from DOB: {extracted_age}")

        # ================================================================
        # EXTRACT ETHNICITY
        # ================================================================
        extracted_ethnicity = None
        ethnicity_patterns = [
            r"(?:ETHNICITY|ETHNIC\s*(?:GROUP|ORIGIN)?)[:\s]+([A-Za-z][A-Za-z\s\-\/]+?)(?:\n|$|,)",
            r"\b(White\s*(?:British|Irish|European|Other)?)\b",
            r"\b(Black\s*(?:British|African|Caribbean|Other)?)\b",
            r"\b(Asian\s*(?:British|Indian|Pakistani|Bangladeshi|Chinese|Other)?)\b",
            r"\b(Mixed\s*(?:White\s*(?:and|&)\s*(?:Black\s*(?:Caribbean|African)|Asian))?)\b",
        ]
        for pattern in ethnicity_patterns:
            match = re.search(pattern, document_text, re.IGNORECASE)
            if match:
                ethnicity_val = match.group(1).strip()
                if len(ethnicity_val) > 2:
                    extracted_ethnicity = ethnicity_val.title()
                    print(f"[MOJ-ASR] Extracted ethnicity: {extracted_ethnicity}")
                    break

        # ================================================================
        # STORE EXTRACTED DATA (apply later when widgets exist)
        # ================================================================
        self._extracted_patient_details = {
            'patient_name': extracted_name,
            'patient_dob': extracted_dob,
            'gender': extracted_gender,
            'age': extracted_age,
            'ethnicity': extracted_ethnicity,
            'hospital': extracted_hospital,
            'mha_section': extracted_mha,
            'mha_section_date': extracted_mha_date,
            'nhs_number': extracted_nhs,
            'diagnoses': extracted_diagnoses,
            'treatment_compliance': extracted_treatment_compliance,
        }

        # Summary of what was extracted
        print(f"[MOJ-ASR] === EXTRACTION SUMMARY ===")
        print(f"[MOJ-ASR] Name: {extracted_name or 'NOT FOUND'}")
        print(f"[MOJ-ASR] DOB: {extracted_dob.strftime('%d/%m/%Y') if extracted_dob else 'NOT FOUND'}")
        print(f"[MOJ-ASR] Gender: {extracted_gender or 'NOT FOUND'}")
        print(f"[MOJ-ASR] Age: {extracted_age or 'NOT FOUND'}")
        print(f"[MOJ-ASR] Ethnicity: {extracted_ethnicity or 'NOT FOUND'}")
        print(f"[MOJ-ASR] Hospital: {extracted_hospital or 'NOT FOUND'}")
        print(f"[MOJ-ASR] MHA: {extracted_mha or 'NOT FOUND'}")
        print(f"[MOJ-ASR] MHA Date: {extracted_mha_date.strftime('%d/%m/%Y') if extracted_mha_date else 'NOT FOUND'}")
        print(f"[MOJ-ASR] NHS: {extracted_nhs or 'NOT FOUND'}")
        print(f"[MOJ-ASR] Diagnoses: {extracted_diagnoses or 'NOT FOUND'}")
        print(f"[MOJ-ASR] Data stored in _extracted_patient_details, will apply when widgets ready")

        # Push patient info to shared store for other forms
        from shared_data_store import get_shared_store
        patient_info = {
            "name": extracted_name,
            "dob": extracted_dob,
            "nhs_number": extracted_nhs,
            "gender": extracted_gender,
            "age": extracted_age,
            "ethnicity": extracted_ethnicity,
        }
        if any(patient_info.values()):
            shared_store = get_shared_store()
            shared_store.set_patient_info(patient_info, source="moj_asr")
            print(f"[MOJ-ASR] Pushed patient details to SharedDataStore: {list(k for k,v in patient_info.items() if v)}")

        # Try to apply now (in case widgets already exist)
        self._apply_extracted_patient_details()

    def _apply_extracted_patient_details(self):
        """Apply stored extracted patient details to UI widgets when they exist."""
        if not hasattr(self, '_extracted_patient_details') or not self._extracted_patient_details:
            return

        details = self._extracted_patient_details
        print(f"[MOJ-ASR] Applying extracted data to UI...")

        # Patient Name - ASR uses popup_patient_name
        extracted_name = details.get('patient_name')
        if extracted_name:
            if hasattr(self, 'popup_patient_name') and self.popup_patient_name:
                if not self.popup_patient_name.text().strip():
                    self.popup_patient_name.setText(extracted_name)
                    print(f"[MOJ-ASR] SET popup_patient_name = '{extracted_name}'")

        # DOB - ASR uses popup_dob
        extracted_dob = details.get('patient_dob')
        if extracted_dob:
            if hasattr(self, 'popup_dob') and self.popup_dob:
                from PySide6.QtCore import QDate
                self.popup_dob.setDate(QDate(extracted_dob.year, extracted_dob.month, extracted_dob.day))
                print(f"[MOJ-ASR] SET popup_dob = {extracted_dob.strftime('%d/%m/%Y')}")

        # Gender - ASR uses popup_gender_male/popup_gender_female
        extracted_gender = details.get('gender')
        if extracted_gender:
            if extracted_gender == "Male":
                if hasattr(self, 'popup_gender_male') and self.popup_gender_male:
                    self.popup_gender_male.setChecked(True)
                    print(f"[MOJ-ASR] SET popup_gender_male = True")
            elif extracted_gender == "Female":
                if hasattr(self, 'popup_gender_female') and self.popup_gender_female:
                    self.popup_gender_female.setChecked(True)
                    print(f"[MOJ-ASR] SET popup_gender_female = True")

        # Hospital - ASR uses popup_hospital
        extracted_hospital = details.get('hospital')
        if extracted_hospital:
            if hasattr(self, 'popup_hospital') and self.popup_hospital:
                if not self.popup_hospital.text().strip():
                    self.popup_hospital.setText(extracted_hospital)
                    print(f"[MOJ-ASR] SET popup_hospital = '{extracted_hospital}'")

        # MHA Section - ASR uses popup_mha_section
        extracted_mha = details.get('mha_section')
        if extracted_mha:
            if hasattr(self, 'popup_mha_section') and self.popup_mha_section:
                for i in range(self.popup_mha_section.count()):
                    if extracted_mha in self.popup_mha_section.itemText(i):
                        self.popup_mha_section.setCurrentIndex(i)
                        print(f"[MOJ-ASR] SET popup_mha_section = '{self.popup_mha_section.itemText(i)}' (index {i})")
                        break

        # MHA Section Date - ASR uses popup_mha_section_date
        extracted_mha_date = details.get('mha_section_date')
        if extracted_mha_date:
            if hasattr(self, 'popup_mha_section_date') and self.popup_mha_section_date:
                from PySide6.QtCore import QDate
                self.popup_mha_section_date.setDate(QDate(extracted_mha_date.year, extracted_mha_date.month, extracted_mha_date.day))
                print(f"[MOJ-ASR] SET popup_mha_section_date = {extracted_mha_date.strftime('%d/%m/%Y')}")

        # NHS Number - ASR uses popup_nhs
        extracted_nhs = details.get('nhs_number')
        if extracted_nhs:
            if hasattr(self, 'popup_nhs') and self.popup_nhs:
                if not self.popup_nhs.text().strip():
                    self.popup_nhs.setText(extracted_nhs)
                    print(f"[MOJ-ASR] SET popup_nhs = '{extracted_nhs}'")

        # Apply diagnoses to Section 3 dx_boxes
        extracted_diagnoses = details.get('diagnoses', [])
        if extracted_diagnoses and hasattr(self, 'dx_boxes') and self.dx_boxes:
            for i, combo in enumerate(self.dx_boxes):
                if i < len(extracted_diagnoses):
                    diag = extracted_diagnoses[i]
                    idx = combo.findText(diag, Qt.MatchFlag.MatchExactly)
                    if idx >= 0:
                        combo.setCurrentIndex(idx)
                    else:
                        idx = combo.findText(diag, Qt.MatchFlag.MatchContains)
                        if idx >= 0:
                            combo.setCurrentIndex(idx)
                        else:
                            combo.setCurrentText(diag)
                    print(f"[MOJ-ASR] Set Section 3 diagnosis {i+1}: {diag}")

        # Apply diagnoses to popup dx_combos
        if extracted_diagnoses and hasattr(self, 'popup_dx_combos') and self.popup_dx_combos:
            for i, combo in enumerate(self.popup_dx_combos):
                if i < len(extracted_diagnoses):
                    diag = extracted_diagnoses[i]
                    idx = combo.findText(diag, Qt.MatchFlag.MatchExactly)
                    if idx >= 0:
                        combo.setCurrentIndex(idx)
                    else:
                        idx = combo.findText(diag, Qt.MatchFlag.MatchContains)
                        if idx >= 0:
                            combo.setCurrentIndex(idx)
                        else:
                            combo.setCurrentText(diag)
                    print(f"[MOJ-ASR] Set popup diagnosis {i+1}: {diag}")

        # Apply Section 6 treatment compliance/understanding
        extracted_treatment_compliance = details.get('treatment_compliance', {})
        if extracted_treatment_compliance and hasattr(self, 'attitude_treatments') and self.attitude_treatments:
            for treatment_key, values in extracted_treatment_compliance.items():
                if treatment_key in self.attitude_treatments:
                    treatment_widgets = self.attitude_treatments[treatment_key]

                    # Set understanding dropdown
                    if "understanding" in values and "understanding" in treatment_widgets:
                        understanding_combo = treatment_widgets["understanding"]
                        understanding_val = values["understanding"]
                        idx = understanding_combo.findText(understanding_val, Qt.MatchFlag.MatchExactly)
                        if idx >= 0:
                            understanding_combo.setCurrentIndex(idx)
                            print(f"[MOJ-ASR] Section 6: Set {treatment_key} understanding = {understanding_val}")

                    # Set compliance dropdown
                    if "compliance" in values and "compliance" in treatment_widgets:
                        compliance_combo = treatment_widgets["compliance"]
                        compliance_val = values["compliance"]
                        idx = compliance_combo.findText(compliance_val, Qt.MatchFlag.MatchExactly)
                        if idx >= 0:
                            compliance_combo.setCurrentIndex(idx)
                            print(f"[MOJ-ASR] Section 6: Set {treatment_key} compliance = {compliance_val}")

    def _create_section_frame(self, title: str, section_num: str = "", color: str = "#991b1b") -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: white; border: none; border-radius: 12px; }")
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)
        display_title = f"{section_num}. {title}" if section_num else title
        title_lbl = QLabel(display_title)
        title_lbl.setStyleSheet(f"font-size: 18px; font-weight: 700; color: {color};")
        layout.addWidget(title_lbl)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 17px;
            }
            QLineEdit:focus { border-color: #991b1b; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 100) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMinimumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
            }
            QTextEdit:focus { border-color: #991b1b; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 17px; }")
        return date_edit

    def _create_prompt_checkbox(self, text: str, color: str = "#0070c0") -> QCheckBox:
        """Create a checkbox for a sub-prompt (blue text from template)."""
        cb = QCheckBox(text)
        cb.setStyleSheet(f"""
            QCheckBox {{
                font-size: 16px;
                color: {color};
                padding: 4px 0;
            }}
            QCheckBox::indicator {{
                width: 16px;
                height: 16px;
            }}
        """)
        return cb

    def _create_prompts_frame(self, prompts: list, color: str = "#0070c0") -> tuple:
        """Create a frame with checkboxes for sub-prompts. Returns (frame, list of checkboxes)."""
        frame = QFrame()
        frame.setStyleSheet(f"QFrame {{ background: #eff6ff; border: none; border-radius: 8px; padding: 8px; }}")
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(12, 8, 12, 8)
        layout.setSpacing(4)

        checkboxes = []
        for prompt_text in prompts:
            cb = self._create_prompt_checkbox(prompt_text, color)
            layout.addWidget(cb)
            checkboxes.append(cb)

        return frame, checkboxes

    # ================================================================
    # SECTION 1: Patient Details (Compact)
    # ================================================================
    def _build_section_1_patient_details(self):
        frame = self._create_section_frame("Patient Details", "1")
        layout = frame.layout()

        # Row 1: Name, DOB, Gender
        row1 = QHBoxLayout()
        row1.setSpacing(8)
        self.patient_name = self._create_line_edit("Patient name (inc. aliases)")
        row1.addWidget(self.patient_name, 3)
        self.patient_dob = self._create_date_edit()
        self.patient_dob.setFixedWidth(120)
        row1.addWidget(self.patient_dob)
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("O")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 15px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
            self.gender_group.addButton(rb)
            row1.addWidget(rb)
        layout.addLayout(row1)

        # Row 2: Hospital, MHCS Ref
        row2 = QHBoxLayout()
        row2.setSpacing(8)
        self.hospital = self._create_line_edit("Detaining Hospital & Address")
        row2.addWidget(self.hospital, 3)
        self.mhcs_ref = self._create_line_edit("MHCS Ref")
        self.mhcs_ref.setFixedWidth(120)
        row2.addWidget(self.mhcs_ref)
        layout.addLayout(row2)

        # Row 3: MHA Section, Section Date, Other Detention
        row3 = QHBoxLayout()
        row3.setSpacing(8)
        self.mha_section = QComboBox()
        self.mha_section.addItems(self.MHA_SECTIONS)
        self.mha_section.setStyleSheet("font-size: 16px; padding: 6px;")
        self.mha_section.setFixedWidth(180)
        row3.addWidget(self.mha_section)
        self.mha_section_date = self._create_date_edit()
        self.mha_section_date.setFixedWidth(120)
        row3.addWidget(self.mha_section_date)
        self.other_detention = self._create_line_edit("Other Detention Authorities (if any)")
        row3.addWidget(self.other_detention, 2)
        layout.addLayout(row3)

        self.form_layout.addWidget(frame)

        # Apply any extracted patient details now that widgets exist
        self._apply_extracted_patient_details()

    # ================================================================
    # SECTION 2: RC Details (Compact)
    # ================================================================
    def _build_section_2_rc_details(self):
        frame = self._create_section_frame("Responsible Clinician's Details", "2", "#059669")
        layout = frame.layout()

        # Row 1: Name, Job Title, Phone
        row1 = QHBoxLayout()
        row1.setSpacing(8)
        self.rc_name = self._create_line_edit("RC Name")
        row1.addWidget(self.rc_name, 2)
        self.rc_job_title = self._create_line_edit("Job Title")
        row1.addWidget(self.rc_job_title, 2)
        self.rc_phone = self._create_line_edit("Phone (direct)")
        row1.addWidget(self.rc_phone, 1)
        layout.addLayout(row1)

        # Row 2: RC Email, MHA Office Email
        row2 = QHBoxLayout()
        row2.setSpacing(8)
        self.rc_email = self._create_line_edit("RC Email (secure)")
        row2.addWidget(self.rc_email, 1)
        self.mha_office_email = self._create_line_edit("MHA Office Email")
        row2.addWidget(self.mha_office_email, 1)
        layout.addLayout(row2)

        self.form_layout.addWidget(frame)

    # ================================================================
    # SECTION 3: Patient's Mental Disorder with Controls Panel (Split Layout)
    # ================================================================
    def _build_section_3_mental_disorder(self):
        frame = self._create_section_frame("Patient's Mental Disorder", "3", "#7c3aed")
        layout = frame.layout()

        info_lbl = QLabel("It is important for the Secretary of State to understand the patient's current mental state and presentation in order to assess the risks they pose to the public.")
        info_lbl.setWordWrap(True)
        info_lbl.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        layout.addWidget(info_lbl)

        split_layout = QHBoxLayout()
        split_layout.setSpacing(16)

        # === LEFT: Mental Disorder Text Area ===
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(8)

        info = QLabel("Click options on the right to auto-generate text:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 15px; color: #6b7280; padding: 6px; background: #f3e8ff; border-radius: 4px;")
        left_layout.addWidget(info)

        self.mental_disorder_text = QTextEdit()
        self.mental_disorder_text.setPlaceholderText("Mental disorder description will be generated here...")
        self.mental_disorder_text.setMinimumHeight(350)
        self.mental_disorder_text.setStyleSheet("""
            QTextEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 12px; font-size: 17px; }
            QTextEdit:focus { border-color: #7c3aed; }
        """)
        left_layout.addWidget(self.mental_disorder_text)

        split_layout.addWidget(left_container, 3)

        # === RIGHT: Controls Panel (scrollable) ===
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        right_scroll.setFixedWidth(380)

        right_container = QWidget()
        right_container.setFixedWidth(360)
        right_container.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 8, 0)
        right_layout.setSpacing(12)

        # --- Mental Disorder (ICD-10) - 3 boxes ---
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(6)

        md_header = QLabel("Mental Disorder")
        md_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        self.dx_boxes = []
        placeholders = ["Primary diagnosis...", "Secondary (optional)...", "Tertiary (optional)..."]
        for i in range(3):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText(placeholders[i])
            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(ICD10_DICT.items(), key=lambda x: x[0].lower()):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})
            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(12)
            combo.setStyleSheet("QComboBox { padding: 5px; font-size: 15px; border: 1px solid #d1d5db; border-radius: 4px; background: white; } QComboBox QAbstractItemView { min-width: 300px; }")
            combo.currentIndexChanged.connect(self._update_mental_disorder_text)
            md_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        self.diagnosis_changed_cb = QCheckBox("Change in diagnosis since last report")
        self.diagnosis_changed_cb.setStyleSheet("font-size: 15px; color: #166534;")
        self.diagnosis_changed_cb.toggled.connect(self._update_mental_disorder_text)
        md_layout.addWidget(self.diagnosis_changed_cb)

        right_layout.addWidget(md_frame)

        # --- Legal Criteria (identical to A8) ---
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(6)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature with sub-options
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._update_mental_disorder_text)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._update_mental_disorder_text)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._update_mental_disorder_text)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree with slider
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(120)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_mental_disorder_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity section
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health with sub-options
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._update_mental_disorder_text)
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._update_mental_disorder_text)
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._update_mental_disorder_text)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety with sub-options
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # === SELF SECTION ===
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 16px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._update_mental_disorder_text)
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Self placement in risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._update_mental_disorder_text)
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._update_mental_disorder_text)
        self_opt_layout.addWidget(self.self_hist_harm)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 16px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._update_mental_disorder_text)
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Self placement in risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._update_mental_disorder_text)
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._update_mental_disorder_text)
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # === OTHERS SECTION ===
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 16px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence to others")
        self.others_hist_violence.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_hist_verbal)

        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_sexual.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_hist_sexual.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_hist_sexual)

        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_stalking.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_hist_stalking.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_hist_stalking)

        self.others_hist_arson = QCheckBox("Arson")
        self.others_hist_arson.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_hist_arson.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_hist_arson)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 16px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence to others")
        self.others_curr_violence.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_sexual.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_curr_sexual.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_curr_sexual)

        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_stalking.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_curr_stalking.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_curr_stalking)

        self.others_curr_arson = QCheckBox("Arson")
        self.others_curr_arson.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_curr_arson.toggled.connect(self._update_mental_disorder_text)
        others_opt_layout.addWidget(self.others_curr_arson)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        right_layout.addWidget(lc_frame)
        right_layout.addStretch()

        right_scroll.setWidget(right_container)
        split_layout.addWidget(right_scroll)

        layout.addLayout(split_layout)
        self.form_layout.addWidget(frame)

        # Apply any extracted patient details (diagnoses) now that widgets exist
        self._apply_extracted_patient_details()

    # --- Control toggle handlers (identical to A8) ---
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_mental_disorder_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_mental_disorder_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_mental_disorder_text()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_mental_disorder_text()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_mental_disorder_text()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_mental_disorder_text()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_mental_disorder_text()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            self.self_hist_neglect.setChecked(False)
            self.self_hist_risky.setChecked(False)
            self.self_hist_harm.setChecked(False)
            self.self_curr_neglect.setChecked(False)
            self.self_curr_risky.setChecked(False)
            self.self_curr_harm.setChecked(False)
        self._update_mental_disorder_text()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            self.others_hist_violence.setChecked(False)
            self.others_hist_verbal.setChecked(False)
            self.others_hist_sexual.setChecked(False)
            self.others_hist_stalking.setChecked(False)
            self.others_hist_arson.setChecked(False)
            self.others_curr_violence.setChecked(False)
            self.others_curr_verbal.setChecked(False)
            self.others_curr_sexual.setChecked(False)
            self.others_curr_stalking.setChecked(False)
            self.others_curr_arson.setChecked(False)
        self._update_mental_disorder_text()

    def _update_mental_disorder_text(self):
        """Update mental disorder text based on control selections, preserving user additions."""
        new_generated = self._generate_mental_disorder_text()
        self._update_text_preserving_additions(self.mental_disorder_text, new_generated, "mental_disorder")

    def _generate_mental_disorder_text(self) -> str:
        """Generate mental disorder text from form selections (identical output style to A8)."""
        p = self._get_pronouns()
        patient = self.patient_name.text().strip() or "The patient"
        paragraphs = []

        # === PARAGRAPH 1: Diagnosis ===
        diagnoses = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    diagnoses.append(f"{dx} ({icd})" if icd else dx)

        if diagnoses:
            if len(diagnoses) == 1:
                para1 = f"{patient} suffers from {diagnoses[0]}."
            else:
                joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                para1 = f"{patient} suffers from {joined}."

            if self.diagnosis_changed_cb.isChecked():
                para1 += f" There has been a change in diagnosis since the last report."

            paragraphs.append(para1)

        # === PARAGRAPH 2: Nature and Degree ===
        nature_degree_parts = []
        if self.nature_cb.isChecked():
            nature_items = []
            if self.relapsing_cb.isChecked():
                nature_items.append("relapsing and remitting")
            if self.treatment_resistant_cb.isChecked():
                nature_items.append("treatment resistant")
            if self.chronic_cb.isChecked():
                nature_items.append("chronic and enduring")

            if nature_items:
                nature_str = ", ".join(nature_items)
                nature_degree_parts.append(f"The nature of {p['pos_l']} mental disorder is {nature_str}.")
            else:
                nature_degree_parts.append(f"The nature of {p['pos_l']} mental disorder warrants continued detention.")

        if self.degree_cb.isChecked():
            level = self.degree_level_label.text().lower()
            symptoms = self.degree_details.text().strip()
            if symptoms:
                nature_degree_parts.append(f"The degree of {p['pos_l']} mental disorder is evidenced by {level} symptoms including {symptoms}.")
            else:
                nature_degree_parts.append(f"The degree of {p['pos_l']} mental disorder is evidenced by {level} symptoms.")

        if nature_degree_parts:
            paragraphs.append(" ".join(nature_degree_parts))

        # === PARAGRAPH 3: Necessity (Health) ===
        health_parts = []
        if self.health_cb.isChecked():
            if self.mental_health_cb.isChecked():
                mh_reasons = []
                if self.poor_compliance_cb.isChecked():
                    mh_reasons.append("poor compliance with treatment")
                if self.limited_insight_cb.isChecked():
                    mh_reasons.append("limited insight into illness")
                if mh_reasons:
                    health_parts.append(f"mental health ({' and '.join(mh_reasons)})")
                else:
                    health_parts.append("mental health")

            if self.physical_health_cb.isChecked():
                ph_details = self.physical_health_details.text().strip()
                if ph_details:
                    health_parts.append(f"physical health ({ph_details})")
                else:
                    health_parts.append("physical health")

            if health_parts:
                paragraphs.append(f"Continued detention is necessary for {p['pos_l']} {' and '.join(health_parts)}.")

        # === PARAGRAPH 4: Necessity (Safety) ===
        if self.safety_cb.isChecked():
            safety_parts = []

            # Self risks
            if self.self_harm_cb.isChecked():
                self_hist = []
                self_curr = []
                if self.self_hist_neglect.isChecked():
                    self_hist.append("self neglect")
                if self.self_hist_risky.isChecked():
                    self_hist.append("self placement in risky situations")
                if self.self_hist_harm.isChecked():
                    self_hist.append("self harm")
                if self.self_curr_neglect.isChecked():
                    self_curr.append("self neglect")
                if self.self_curr_risky.isChecked():
                    self_curr.append("self placement in risky situations")
                if self.self_curr_harm.isChecked():
                    self_curr.append("self harm")

                self_parts = []
                if self_hist:
                    self_parts.append(f"historically {', '.join(self_hist)}")
                if self_curr:
                    self_parts.append(f"currently {', '.join(self_curr)}")
                if self_parts:
                    safety_parts.append(f"risk to {p['self']} ({'; '.join(self_parts)})")
                else:
                    safety_parts.append(f"risk to {p['self']}")

            # Others risks
            if self.others_cb.isChecked():
                others_hist = []
                others_curr = []
                if self.others_hist_violence.isChecked():
                    others_hist.append("violence")
                if self.others_hist_verbal.isChecked():
                    others_hist.append("verbal aggression")
                if self.others_hist_sexual.isChecked():
                    others_hist.append("sexual violence")
                if self.others_hist_stalking.isChecked():
                    others_hist.append("stalking")
                if self.others_hist_arson.isChecked():
                    others_hist.append("arson")
                if self.others_curr_violence.isChecked():
                    others_curr.append("violence")
                if self.others_curr_verbal.isChecked():
                    others_curr.append("verbal aggression")
                if self.others_curr_sexual.isChecked():
                    others_curr.append("sexual violence")
                if self.others_curr_stalking.isChecked():
                    others_curr.append("stalking")
                if self.others_curr_arson.isChecked():
                    others_curr.append("arson")

                others_parts = []
                if others_hist:
                    others_parts.append(f"historically {', '.join(others_hist)}")
                if others_curr:
                    others_parts.append(f"currently {', '.join(others_curr)}")
                if others_parts:
                    safety_parts.append(f"risk to others ({'; '.join(others_parts)})")
                else:
                    safety_parts.append("risk to others")

            if safety_parts:
                paragraphs.append(f"Continued detention is necessary for the protection of {p['obj']} and others due to {' and '.join(safety_parts)}.")

        return "\n\n".join(paragraphs)

    # ================================================================
    # SECTION 4: Attitude and Behaviour
    # ================================================================
    def _build_section_4_attitude_behaviour(self):
        frame = self._create_section_frame("Attitude and Behaviour", "4", "#7c3aed")
        layout = frame.layout()

        prompt_lbl = QLabel("Please describe the patient's attitude and behaviour in the last 12 months, and/or the community, including any concerns:")
        prompt_lbl.setWordWrap(True)
        prompt_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(prompt_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.attitude_behaviour_text = self._create_text_edit("", 300)
        self.attitude_behaviour_text.setMinimumWidth(350)
        split.addWidget(self.attitude_behaviour_text, 2)

        # Right panel container
        right_container = QWidget()
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(8)

        # === BEHAVIOUR CATEGORIES - Collapsible Section ===
        self.behaviour_section = CollapsibleSection("Behaviour Categories") if CollapsibleSection else None

        # Container for behaviour items grid
        behaviour_content = QWidget()
        behaviour_grid = QGridLayout(behaviour_content)
        behaviour_grid.setContentsMargins(8, 8, 8, 8)
        behaviour_grid.setSpacing(6)

        self.behaviour_items = {}
        behaviour_categories = [
            ("verbal_aggression", "Verbal aggression", "no verbal aggression noted"),
            ("violence_people", "Violence to people", "no violence to people noted"),
            ("violence_objects", "Violence to objects", "no violence to objects noted"),
            ("substance_abuse", "Substance abuse", "no substance abuse noted"),
            ("self_harm", "Self-harm", "no self-harm noted"),
            ("fire_setting", "Fire-setting", "no fire-setting noted"),
            ("secretive", "Manipulative behaviour", "no manipulative behaviour noted"),
            ("subversive", "Subversive behaviour", "no subversive behaviour noted"),
            ("sexually_disinhibited", "Sexually disinhibited", "no sexually disinhibited behaviour noted"),
            ("inappropriate", "Inappropriate behaviour", "no inappropriate behaviour noted"),
            ("antisocial", "Antisocial behaviour", "no antisocial behaviour noted"),
            ("extremist", "Extremist behaviour", "no extremist behaviour noted"),
            ("seclusion", "Periods of seclusion", "no periods of seclusion noted")
        ]

        row = 0
        col = 0
        for key, label, negative_text in behaviour_categories:
            # Compact item: Label | Yes | No in one row
            item_widget = QWidget()
            item_layout = QHBoxLayout(item_widget)
            item_layout.setContentsMargins(4, 2, 4, 2)
            item_layout.setSpacing(4)

            lbl = QLabel(label)
            lbl.setStyleSheet("font-size: 15px; color: #374151;")
            lbl.setFixedWidth(130)
            item_layout.addWidget(lbl)

            yes_rb = QRadioButton("Y")
            yes_rb.setStyleSheet("""
                QRadioButton {
                    font-size: 14px;
                    color: #059669;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 14px;
                    height: 14px;
                }
            """)
            yes_rb.setFixedWidth(32)
            no_rb = QRadioButton("N")
            no_rb.setStyleSheet("""
                QRadioButton {
                    font-size: 14px;
                    color: #dc2626;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 14px;
                    height: 14px;
                }
            """)
            no_rb.setFixedWidth(32)

            btn_group = QButtonGroup(item_widget)
            btn_group.addButton(yes_rb, 1)
            btn_group.addButton(no_rb, 0)

            item_layout.addWidget(yes_rb)
            item_layout.addWidget(no_rb)

            # Details input (hidden initially)
            details_input = QLineEdit()
            details_input.setPlaceholderText("Details...")
            details_input.setStyleSheet("font-size: 14px; padding: 2px 4px; border: 1px solid #d1d5db; border-radius: 2px;")
            details_input.setFixedWidth(100)
            details_input.hide()
            details_input.textChanged.connect(self._update_behaviour_text)
            item_layout.addWidget(details_input)

            yes_rb.toggled.connect(lambda checked, d=details_input: d.setVisible(checked))
            yes_rb.toggled.connect(self._update_behaviour_text)
            no_rb.toggled.connect(self._update_behaviour_text)

            behaviour_grid.addWidget(item_widget, row, col)

            self.behaviour_items[key] = {
                "yes": yes_rb,
                "no": no_rb,
                "details": details_input,
                "label": label,
                "negative": negative_text
            }

            col += 1
            if col >= 2:  # 2 items per row
                col = 0
                row += 1

        if self.behaviour_section:
            self.behaviour_section.set_content(behaviour_content)
            self.behaviour_section.set_content_height(220)
            right_layout.addWidget(self.behaviour_section)
        else:
            right_layout.addWidget(behaviour_content)

        # === IMPORTED DATA - Collapsible Section ===
        self.import_section = CollapsibleSection("Imported Data", start_collapsed=True) if CollapsibleSection else None

        self.behaviour_import_container = QWidget()
        self.behaviour_import_layout = QVBoxLayout(self.behaviour_import_container)
        self.behaviour_import_layout.setContentsMargins(8, 8, 8, 8)
        self.behaviour_import_layout.setSpacing(4)

        self.behaviour_import_placeholder = QLabel("No imported data. Use Import File to upload data.")
        self.behaviour_import_placeholder.setStyleSheet("color: #64748b; font-style: italic; font-size: 15px;")
        self.behaviour_import_layout.addWidget(self.behaviour_import_placeholder)

        self.behaviour_imported_entries = []

        if self.import_section:
            self.import_section.set_content(self.behaviour_import_container)
            self.import_section.set_content_height(150)
            right_layout.addWidget(self.import_section)
        else:
            # Fallback without CollapsibleSection
            import_frame = QFrame()
            import_frame.setStyleSheet("QFrame { background: #f0f9ff; border: none; border-radius: 6px; }")
            import_inner = QVBoxLayout(import_frame)
            import_inner.setContentsMargins(8, 8, 8, 8)
            import_title = QLabel("Imported Data")
            import_title.setStyleSheet("font-size: 16px; font-weight: 700; color: #0369a1;")
            import_inner.addWidget(import_title)
            import_inner.addWidget(self.behaviour_import_container)
            right_layout.addWidget(import_frame)

        right_layout.addStretch()
        split.addWidget(right_container, 3)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def populate_behaviour_import_data(self, entries: list):
        """Populate the imported data panel with entries that have checkboxes."""
        # Check if layout exists (section 4 may use popup version instead)
        if not hasattr(self, 'behaviour_import_layout') or not self.behaviour_import_layout:
            print("[MOJ-ASR] Section 4: behaviour_import_layout not available, skipping populate_behaviour_import_data")
            return

        # Clear existing entries
        while self.behaviour_import_layout.count():
            item = self.behaviour_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.behaviour_imported_entries = []

        if not entries:
            self.behaviour_import_placeholder = QLabel("No imported data. Use Import File to upload data.")
            self.behaviour_import_placeholder.setStyleSheet("color: #64748b; font-style: italic; font-size: 15px;")
            self.behaviour_import_layout.addWidget(self.behaviour_import_placeholder)
            return

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date = entry.get("date", "") or entry.get("datetime", "")
            if not text:
                continue
            # Remove blank lines
            text = '\n'.join(line for line in text.split('\n') if line.strip())

            entry_frame = QFrame()
            entry_frame.setStyleSheet("QFrame { background: white; border: none; border-radius: 4px; }")
            entry_layout = QHBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 4, 6, 4)
            entry_layout.setSpacing(6)

            cb = QCheckBox()
            cb.setStyleSheet("QCheckBox { margin-right: 2px; }")
            entry_layout.addWidget(cb)

            content_layout = QVBoxLayout()
            content_layout.setSpacing(1)

            if date:
                date_lbl = QLabel(str(date))
                date_lbl.setStyleSheet("font-size: 22px; font-weight: 600; color: #64748b;")
                content_layout.addWidget(date_lbl)

            # Truncate text for display
            display_text = text[:150] + "..." if len(text) > 150 else text
            text_lbl = QLabel(display_text)
            text_lbl.setWordWrap(True)
            text_lbl.setStyleSheet("font-size: 16px; color: #374151;")
            content_layout.addWidget(text_lbl)

            entry_layout.addLayout(content_layout, 1)
            self.behaviour_import_layout.addWidget(entry_frame)

            # Store checkbox with full text
            entry_data = {"checkbox": cb, "text": text, "date": date}
            self.behaviour_imported_entries.append(entry_data)

            # Connect checkbox to update text
            cb.toggled.connect(self._on_behaviour_import_toggled)

        # Expand the import section if it exists and has CollapsibleSection
        if self.import_section and hasattr(self.import_section, '_toggle_collapse'):
            if self.import_section.is_collapsed():
                self.import_section._toggle_collapse()

    def _on_behaviour_import_toggled(self):
        """Handle toggling of imported behaviour entries - add/remove from text area."""
        current_text = self.attitude_behaviour_text.toPlainText()

        # Collect all checked entries
        checked_texts = []
        for entry in self.behaviour_imported_entries:
            if entry["checkbox"].isChecked():
                date = entry.get("date", "")
                text = entry["text"]
                if date:
                    checked_texts.append(f"[{date}] {text}")
                else:
                    checked_texts.append(text)

        # Find the import section marker or append
        import_marker = "\n\n--- Imported Notes ---\n"
        if import_marker in current_text:
            # Replace the imported section
            base_text = current_text.split(import_marker)[0]
        else:
            base_text = current_text

        if checked_texts:
            new_text = base_text.rstrip() + import_marker + "\n".join(checked_texts)
        else:
            new_text = base_text.rstrip()

        self.attitude_behaviour_text.setPlainText(new_text)

    def _update_behaviour_text(self):
        """Update behaviour text based on Yes/No selections and details."""
        concerns = []
        negatives = []

        for key, item in self.behaviour_items.items():
            if item["yes"].isChecked():
                label = item["label"].lower()
                details = item["details"].text().strip()
                if details:
                    concerns.append(f"{label} ({details})")
                else:
                    concerns.append(label)
            elif item["no"].isChecked():
                negatives.append(item["negative"])

        parts = []

        # Positive concerns
        if concerns:
            if len(concerns) == 1:
                parts.append(f"In the last 12 months there have been concerns regarding {concerns[0]}.")
            else:
                items = ", ".join(concerns[:-1]) + f" and {concerns[-1]}"
                parts.append(f"In the last 12 months there have been concerns regarding {items}.")

        # Negative statements
        if negatives:
            if len(negatives) == 1:
                parts.append(f"There has been {negatives[0]}.")
            else:
                neg_items = ", ".join(negatives[:-1]) + f" and {negatives[-1]}"
                parts.append(f"There has been {neg_items}.")

        if parts:
            new_generated = " ".join(parts)
            self._update_text_preserving_additions(self.attitude_behaviour_text, new_generated, "behaviour")

    # ================================================================
    # SECTION 5: Addressing Issues
    # ================================================================
    def _build_section_5_addressing_issues(self):
        frame = self._create_section_frame("Addressing Issues", "5", "#7c3aed")
        layout = frame.layout()

        prompt_lbl = QLabel("Please state what effect these have had on the patient and how they have been addressed:")
        prompt_lbl.setWordWrap(True)
        prompt_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(prompt_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        # Left side: text area
        self.addressing_issues_text = self._create_text_edit("", 400)
        self.addressing_issues_text.setMinimumWidth(350)
        split.addWidget(self.addressing_issues_text, 2)

        # Right side: scrollable controls panel
        right_container = QWidget()
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(8)

        # === PSYCHOLOGY WORK SECTION (Collapsible) ===
        self.psych_work_section = CollapsibleSection("Psychology Work") if CollapsibleSection else None

        psych_content = QWidget()
        psych_layout = QVBoxLayout(psych_content)
        psych_layout.setContentsMargins(12, 12, 12, 12)
        psych_layout.setSpacing(12)

        # Index Offence Work slider
        idx_frame = QFrame()
        idx_frame.setStyleSheet("QFrame { background: #f9fafb; border-radius: 6px; }")
        idx_layout = QVBoxLayout(idx_frame)
        idx_layout.setContentsMargins(10, 10, 10, 10)
        idx_layout.setSpacing(6)

        idx_lbl = QLabel("Index Offence Work:")
        idx_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1f2937;")
        idx_layout.addWidget(idx_lbl)

        # Progress slider row
        idx_progress_row = QHBoxLayout()
        idx_progress_lbl = QLabel("Progress:")
        idx_progress_lbl.setStyleSheet("font-size: 15px; color: #4b5563; min-width: 60px;")
        idx_progress_row.addWidget(idx_progress_lbl)

        self.index_offence_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.index_offence_slider.setMinimum(0)
        self.index_offence_slider.setMaximum(4)
        self.index_offence_slider.setValue(0)
        self.index_offence_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.index_offence_slider.setTickInterval(1)
        self.index_offence_slider.setFixedWidth(160)
        idx_progress_row.addWidget(self.index_offence_slider)

        self.index_offence_level = QLabel("No")
        self.index_offence_level.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937; min-width: 70px;")
        idx_progress_row.addWidget(self.index_offence_level)
        idx_progress_row.addStretch()
        idx_layout.addLayout(idx_progress_row)

        self.index_offence_slider.valueChanged.connect(lambda v: self.index_offence_level.setText(["No", "Started", "Ongoing", "Advanced", "Completed"][v]))
        self.index_offence_slider.valueChanged.connect(self._update_addressing_text)

        # Effectiveness slider row
        idx_eff_row = QHBoxLayout()
        idx_eff_lbl = QLabel("Effectiveness:")
        idx_eff_lbl.setStyleSheet("font-size: 15px; color: #4b5563; min-width: 60px;")
        idx_eff_row.addWidget(idx_eff_lbl)

        self.index_offence_effectiveness = NoWheelSlider(Qt.Orientation.Horizontal)
        self.index_offence_effectiveness.setMinimum(0)
        self.index_offence_effectiveness.setMaximum(4)
        self.index_offence_effectiveness.setValue(0)
        self.index_offence_effectiveness.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.index_offence_effectiveness.setTickInterval(1)
        self.index_offence_effectiveness.setFixedWidth(160)
        idx_eff_row.addWidget(self.index_offence_effectiveness)

        self.index_offence_eff_level = QLabel("No")
        self.index_offence_eff_level.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937; min-width: 70px;")
        idx_eff_row.addWidget(self.index_offence_eff_level)
        idx_eff_row.addStretch()
        idx_layout.addLayout(idx_eff_row)

        self.index_offence_effectiveness.valueChanged.connect(lambda v: self.index_offence_eff_level.setText(["No", "Some", "Moderate", "Significant", "High"][v]))
        self.index_offence_effectiveness.valueChanged.connect(self._update_addressing_text)

        self.index_offence_details = QLineEdit()
        self.index_offence_details.setPlaceholderText("Details...")
        self.index_offence_details.setStyleSheet("font-size: 15px; padding: 4px 6px;")
        self.index_offence_details.textChanged.connect(self._update_addressing_text)
        idx_layout.addWidget(self.index_offence_details)

        psych_layout.addWidget(idx_frame)

        # Risks Work slider
        risk_frame = QFrame()
        risk_frame.setStyleSheet("QFrame { background: #f9fafb; border-radius: 6px; }")
        risk_layout = QVBoxLayout(risk_frame)
        risk_layout.setContentsMargins(10, 10, 10, 10)
        risk_layout.setSpacing(6)

        risk_lbl = QLabel("Risks Work:")
        risk_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1f2937;")
        risk_layout.addWidget(risk_lbl)

        # Progress slider row
        risk_progress_row = QHBoxLayout()
        risk_progress_lbl = QLabel("Progress:")
        risk_progress_lbl.setStyleSheet("font-size: 15px; color: #4b5563; min-width: 60px;")
        risk_progress_row.addWidget(risk_progress_lbl)

        self.risks_work_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.risks_work_slider.setMinimum(0)
        self.risks_work_slider.setMaximum(4)
        self.risks_work_slider.setValue(0)
        self.risks_work_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.risks_work_slider.setTickInterval(1)
        self.risks_work_slider.setFixedWidth(160)
        risk_progress_row.addWidget(self.risks_work_slider)

        self.risks_work_level = QLabel("No")
        self.risks_work_level.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937; min-width: 70px;")
        risk_progress_row.addWidget(self.risks_work_level)
        risk_progress_row.addStretch()
        risk_layout.addLayout(risk_progress_row)

        self.risks_work_slider.valueChanged.connect(lambda v: self.risks_work_level.setText(["No", "Started", "Ongoing", "Advanced", "Completed"][v]))
        self.risks_work_slider.valueChanged.connect(self._update_addressing_text)

        # Effectiveness slider row
        risk_eff_row = QHBoxLayout()
        risk_eff_lbl = QLabel("Effectiveness:")
        risk_eff_lbl.setStyleSheet("font-size: 15px; color: #4b5563; min-width: 60px;")
        risk_eff_row.addWidget(risk_eff_lbl)

        self.risks_work_effectiveness = NoWheelSlider(Qt.Orientation.Horizontal)
        self.risks_work_effectiveness.setMinimum(0)
        self.risks_work_effectiveness.setMaximum(4)
        self.risks_work_effectiveness.setValue(0)
        self.risks_work_effectiveness.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.risks_work_effectiveness.setTickInterval(1)
        self.risks_work_effectiveness.setFixedWidth(160)
        risk_eff_row.addWidget(self.risks_work_effectiveness)

        self.risks_work_eff_level = QLabel("No")
        self.risks_work_eff_level.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937; min-width: 70px;")
        risk_eff_row.addWidget(self.risks_work_eff_level)
        risk_eff_row.addStretch()
        risk_layout.addLayout(risk_eff_row)

        self.risks_work_effectiveness.valueChanged.connect(lambda v: self.risks_work_eff_level.setText(["No", "Some", "Moderate", "Significant", "High"][v]))
        self.risks_work_effectiveness.valueChanged.connect(self._update_addressing_text)

        self.risks_work_details = QLineEdit()
        self.risks_work_details.setPlaceholderText("Details...")
        self.risks_work_details.setStyleSheet("font-size: 15px; padding: 4px 6px;")
        self.risks_work_details.textChanged.connect(self._update_addressing_text)
        risk_layout.addWidget(self.risks_work_details)

        psych_layout.addWidget(risk_frame)

        if self.psych_work_section:
            self.psych_work_section.set_content(psych_content)
            self.psych_work_section.set_content_height(380)
            right_layout.addWidget(self.psych_work_section)
        else:
            right_layout.addWidget(psych_content)

        # === WARD ACTIVITIES SECTION (Collapsible) ===
        self.ward_activities_section = CollapsibleSection("Ward Activities") if CollapsibleSection else None

        ward_content = QWidget()
        ward_layout = QVBoxLayout(ward_content)
        ward_layout.setContentsMargins(12, 12, 12, 12)
        ward_layout.setSpacing(8)

        # OT checkbox with sub-options
        self.ot_cb = QCheckBox("OT")
        self.ot_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        ward_layout.addWidget(self.ot_cb)

        self.ot_options = QWidget()
        ot_options_layout = QVBoxLayout(self.ot_options)
        ot_options_layout.setContentsMargins(20, 4, 0, 8)
        ot_options_layout.setSpacing(6)

        # OT activity checkboxes grid
        ot_grid_widget = QWidget()
        ot_grid = QGridLayout(ot_grid_widget)
        ot_grid.setContentsMargins(0, 0, 0, 0)
        ot_grid.setSpacing(4)

        ot_items = ["Education", "Cooking groups", "OT outings", "Breakfast club", "Smoothie group",
                    "Sports/gym", "Current affairs", "Music group", "Arts group", "Self care", "Budgeting"]
        self.ot_checkboxes = {}
        for i, item in enumerate(ot_items):
            cb = QCheckBox(item)
            cb.setStyleSheet("font-size: 15px; color: #4b5563;")
            cb.toggled.connect(self._update_addressing_text)
            ot_grid.addWidget(cb, i // 3, i % 3)
            self.ot_checkboxes[item.lower().replace(" ", "_").replace("/", "_")] = cb
        ot_options_layout.addWidget(ot_grid_widget)

        # OT Concerns slider
        ot_concerns_row = QHBoxLayout()
        ot_concerns_lbl = QLabel("Concerns:")
        ot_concerns_lbl.setStyleSheet("font-size: 15px; color: #4b5563; min-width: 60px;")
        ot_concerns_row.addWidget(ot_concerns_lbl)

        self.ot_concerns_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.ot_concerns_slider.setMinimum(0)
        self.ot_concerns_slider.setMaximum(4)
        self.ot_concerns_slider.setValue(0)
        self.ot_concerns_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.ot_concerns_slider.setTickInterval(1)
        self.ot_concerns_slider.setFixedWidth(140)
        ot_concerns_row.addWidget(self.ot_concerns_slider)

        self.ot_concerns_level = QLabel("No")
        self.ot_concerns_level.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937; min-width: 70px;")
        ot_concerns_row.addWidget(self.ot_concerns_level)
        ot_concerns_row.addStretch()
        ot_options_layout.addLayout(ot_concerns_row)

        self.ot_concerns_slider.valueChanged.connect(lambda v: self.ot_concerns_level.setText(["No", "Minor", "Some", "Moderate", "Significant"][v]))
        self.ot_concerns_slider.valueChanged.connect(self._update_ot_manageable_visibility)
        self.ot_concerns_slider.valueChanged.connect(self._update_addressing_text)

        # OT Manageable radios (hidden when no concerns)
        self.ot_manageable_container = QWidget()
        ot_manage_layout = QHBoxLayout(self.ot_manageable_container)
        ot_manage_layout.setContentsMargins(0, 0, 0, 0)
        ot_manage_layout.setSpacing(12)

        self.ot_manageable_group = QButtonGroup(self.ot_manageable_container)
        self.ot_manageable = QRadioButton("Manageable")
        self.ot_manageable.setStyleSheet("font-size: 15px; color: #059669;")
        self.ot_not_manageable = QRadioButton("Not Manageable")
        self.ot_not_manageable.setStyleSheet("font-size: 15px; color: #dc2626;")
        self.ot_manageable_group.addButton(self.ot_manageable)
        self.ot_manageable_group.addButton(self.ot_not_manageable)
        ot_manage_layout.addWidget(self.ot_manageable)
        ot_manage_layout.addWidget(self.ot_not_manageable)
        ot_manage_layout.addStretch()

        self.ot_manageable.toggled.connect(self._update_addressing_text)
        self.ot_not_manageable.toggled.connect(self._update_addressing_text)
        self.ot_manageable_container.hide()
        ot_options_layout.addWidget(self.ot_manageable_container)

        self.ot_options.hide()
        self.ot_cb.toggled.connect(lambda checked: self.ot_options.setVisible(checked))
        self.ot_cb.toggled.connect(self._update_addressing_text)
        ward_layout.addWidget(self.ot_options)

        # Psychology checkbox with sub-options
        self.psych_groups_cb = QCheckBox("Psychology Groups")
        self.psych_groups_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        ward_layout.addWidget(self.psych_groups_cb)

        self.psych_options = QWidget()
        psych_options_layout = QVBoxLayout(self.psych_options)
        psych_options_layout.setContentsMargins(20, 4, 0, 8)
        psych_options_layout.setSpacing(6)

        # Psychology activity checkboxes grid
        psych_grid_widget = QWidget()
        psych_grid = QGridLayout(psych_grid_widget)
        psych_grid.setContentsMargins(0, 0, 0, 0)
        psych_grid.setSpacing(4)

        psych_items = ["Psychoeducation", "Understanding mental illness", "Emotional regulation",
                       "Social behaviour", "Understanding psychosis", "Understanding risk",
                       "Discharge planning", "Drugs and alcohol"]
        self.psych_checkboxes = {}
        for i, item in enumerate(psych_items):
            cb = QCheckBox(item)
            cb.setStyleSheet("font-size: 15px; color: #4b5563;")
            cb.toggled.connect(self._update_addressing_text)
            psych_grid.addWidget(cb, i // 2, i % 2)
            self.psych_checkboxes[item.lower().replace(" ", "_")] = cb
        psych_options_layout.addWidget(psych_grid_widget)

        # Psychology Concerns slider
        psych_concerns_row = QHBoxLayout()
        psych_concerns_lbl = QLabel("Concerns:")
        psych_concerns_lbl.setStyleSheet("font-size: 15px; color: #4b5563; min-width: 60px;")
        psych_concerns_row.addWidget(psych_concerns_lbl)

        self.psych_concerns_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.psych_concerns_slider.setMinimum(0)
        self.psych_concerns_slider.setMaximum(4)
        self.psych_concerns_slider.setValue(0)
        self.psych_concerns_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.psych_concerns_slider.setTickInterval(1)
        self.psych_concerns_slider.setFixedWidth(140)
        psych_concerns_row.addWidget(self.psych_concerns_slider)

        self.psych_concerns_level = QLabel("No")
        self.psych_concerns_level.setStyleSheet("font-size: 15px; font-weight: 600; color: #1f2937; min-width: 70px;")
        psych_concerns_row.addWidget(self.psych_concerns_level)
        psych_concerns_row.addStretch()
        psych_options_layout.addLayout(psych_concerns_row)

        self.psych_concerns_slider.valueChanged.connect(lambda v: self.psych_concerns_level.setText(["No", "Minor", "Some", "Moderate", "Significant"][v]))
        self.psych_concerns_slider.valueChanged.connect(self._update_psych_manageable_visibility)
        self.psych_concerns_slider.valueChanged.connect(self._update_addressing_text)

        # Psychology Manageable radios (hidden when no concerns)
        self.psych_manageable_container = QWidget()
        psych_manage_layout = QHBoxLayout(self.psych_manageable_container)
        psych_manage_layout.setContentsMargins(0, 0, 0, 0)
        psych_manage_layout.setSpacing(12)

        self.psych_manageable_group = QButtonGroup(self.psych_manageable_container)
        self.psych_manageable = QRadioButton("Manageable")
        self.psych_manageable.setStyleSheet("font-size: 15px; color: #059669;")
        self.psych_not_manageable = QRadioButton("Not Manageable")
        self.psych_not_manageable.setStyleSheet("font-size: 15px; color: #dc2626;")
        self.psych_manageable_group.addButton(self.psych_manageable)
        self.psych_manageable_group.addButton(self.psych_not_manageable)
        psych_manage_layout.addWidget(self.psych_manageable)
        psych_manage_layout.addWidget(self.psych_not_manageable)
        psych_manage_layout.addStretch()

        self.psych_manageable.toggled.connect(self._update_addressing_text)
        self.psych_not_manageable.toggled.connect(self._update_addressing_text)
        self.psych_manageable_container.hide()
        psych_options_layout.addWidget(self.psych_manageable_container)

        self.psych_options.hide()
        self.psych_groups_cb.toggled.connect(lambda checked: self.psych_options.setVisible(checked))
        self.psych_groups_cb.toggled.connect(self._update_addressing_text)
        ward_layout.addWidget(self.psych_options)

        if self.ward_activities_section:
            self.ward_activities_section.set_content(ward_content)
            self.ward_activities_section.set_content_height(320)
            right_layout.addWidget(self.ward_activities_section)
        else:
            right_layout.addWidget(ward_content)

        # === CURRENT UNDERSTANDING SECTION (Collapsible) ===
        self.understanding_section = CollapsibleSection("Current Understanding of Risks") if CollapsibleSection else None

        understanding_content = QWidget()
        understanding_layout = QVBoxLayout(understanding_content)
        understanding_layout.setContentsMargins(12, 12, 12, 12)
        understanding_layout.setSpacing(6)

        # Risk items with Understanding dropdown and Engagement slider (matches Section 5)
        RISK_TYPES = [
            ("violence_others", "Violence to others"),
            ("violence_property", "Violence to property"),
            ("verbal_aggression", "Verbal aggression"),
            ("substance_misuse", "Substance misuse"),
            ("self_harm", "Self harm"),
            ("self_neglect", "Self neglect"),
            ("stalking", "Stalking"),
            ("threatening_behaviour", "Threatening behaviour"),
            ("sexually_inappropriate", "Sexually inappropriate behaviour"),
            ("vulnerability", "Vulnerability"),
            ("bullying_victimisation", "Bullying/victimisation"),
            ("absconding", "Absconding/AWOL"),
            ("reoffending", "Reoffending"),
        ]

        self.risk_understanding = {}
        for key, label in RISK_TYPES:
            risk_row = QFrame()
            risk_row.setStyleSheet("QFrame { background: #f9fafb; border-radius: 4px; padding: 4px; }")
            row_layout = QVBoxLayout(risk_row)
            row_layout.setContentsMargins(8, 6, 8, 6)
            row_layout.setSpacing(4)

            # Checkbox for the risk
            risk_cb = QCheckBox(label)
            risk_cb.setStyleSheet("font-size: 15px; font-weight: 600; color: #374151;")
            row_layout.addWidget(risk_cb)

            # Hidden controls
            controls = QWidget()
            controls_layout = QHBoxLayout(controls)
            controls_layout.setContentsMargins(16, 4, 0, 0)
            controls_layout.setSpacing(8)

            # Understanding dropdown
            und_lbl = QLabel("Understanding:")
            und_lbl.setStyleSheet("font-size: 14px; color: #6b7280;")
            controls_layout.addWidget(und_lbl)

            und_combo = QComboBox()
            und_combo.addItems(["Select...", "Good", "Fair", "Poor"])
            und_combo.setStyleSheet("font-size: 14px;")
            und_combo.setFixedWidth(80)
            und_combo.currentTextChanged.connect(self._update_addressing_text)
            controls_layout.addWidget(und_combo)

            # Engagement slider
            eng_lbl = QLabel("Engagement:")
            eng_lbl.setStyleSheet("font-size: 14px; color: #6b7280;")
            controls_layout.addWidget(eng_lbl)

            eng_slider = NoWheelSlider(Qt.Orientation.Horizontal)
            eng_slider.setMinimum(0)
            eng_slider.setMaximum(4)
            eng_slider.setValue(0)
            eng_slider.setFixedWidth(100)
            eng_slider.valueChanged.connect(self._update_addressing_text)
            controls_layout.addWidget(eng_slider)

            eng_level = QLabel("No")
            eng_level.setStyleSheet("font-size: 14px; color: #6b7280; min-width: 60px;")
            eng_slider.valueChanged.connect(lambda v, l=eng_level: l.setText(["No", "Started", "Ongoing", "Advanced", "Completed"][v]))
            controls_layout.addWidget(eng_level)
            controls_layout.addStretch()

            controls.hide()
            risk_cb.toggled.connect(lambda checked, c=controls: c.setVisible(checked))
            risk_cb.toggled.connect(self._update_addressing_text)
            row_layout.addWidget(controls)

            understanding_layout.addWidget(risk_row)

            self.risk_understanding[key] = {
                "checkbox": risk_cb,
                "understanding": und_combo,
                "engagement_slider": eng_slider,
                "engagement_label": eng_level
            }

        if self.understanding_section:
            self.understanding_section.set_content(understanding_content)
            self.understanding_section.set_content_height(250)
            right_layout.addWidget(self.understanding_section)
        else:
            right_layout.addWidget(understanding_content)

        # === EFFECTIVENESS SECTION ===
        eff_frame = QFrame()
        eff_frame.setStyleSheet("QFrame { background: #fef3c7; border: none; border-radius: 6px; }")
        eff_layout = QVBoxLayout(eff_frame)
        eff_layout.setContentsMargins(12, 8, 12, 8)
        eff_layout.setSpacing(6)

        eff_lbl = QLabel("Overall Effectiveness:")
        eff_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #806000;")
        eff_layout.addWidget(eff_lbl)

        eff_row = QHBoxLayout()
        self.effectiveness_group = QButtonGroup(eff_frame)
        self.effectiveness_radios = {}
        for level in ["Mild", "Moderate", "High", "Very High"]:
            rb = QRadioButton(level)
            rb.setStyleSheet("font-size: 15px; color: #78350f;")
            rb.toggled.connect(self._update_addressing_text)
            eff_row.addWidget(rb)
            self.effectiveness_group.addButton(rb)
            self.effectiveness_radios[level.lower().replace(" ", "_")] = rb
        eff_row.addStretch()
        eff_layout.addLayout(eff_row)

        right_layout.addWidget(eff_frame)

        # === RELAPSE PREVENTION SECTION ===
        relapse_frame = QFrame()
        relapse_frame.setStyleSheet("QFrame { background: #ecfdf5; border: none; border-radius: 6px; }")
        relapse_layout = QVBoxLayout(relapse_frame)
        relapse_layout.setContentsMargins(12, 8, 12, 8)
        relapse_layout.setSpacing(6)

        relapse_lbl = QLabel("Relapse Prevention:")
        relapse_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #065f46;")
        relapse_layout.addWidget(relapse_lbl)

        relapse_row = QHBoxLayout()
        self.relapse_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.relapse_slider.setMinimum(0)
        self.relapse_slider.setMaximum(4)
        self.relapse_slider.setValue(0)
        self.relapse_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.relapse_slider.setTickInterval(1)
        self.relapse_slider.setFixedWidth(180)
        relapse_row.addWidget(self.relapse_slider)

        self.relapse_level = QLabel("Nil")
        self.relapse_level.setStyleSheet("font-size: 15px; font-weight: 500; color: #047857; min-width: 100px;")
        self.relapse_slider.valueChanged.connect(lambda v: self.relapse_level.setText(["Nil", "Started", "Ongoing", "Almost Completed", "Completed"][v]))
        self.relapse_slider.valueChanged.connect(self._update_addressing_text)
        relapse_row.addWidget(self.relapse_level)

        self.relapse_refused = QCheckBox("Refused")
        self.relapse_refused.setStyleSheet("font-size: 15px; color: #dc2626;")
        self.relapse_refused.toggled.connect(self._update_addressing_text)
        relapse_row.addWidget(self.relapse_refused)
        relapse_row.addStretch()

        relapse_layout.addLayout(relapse_row)
        right_layout.addWidget(relapse_frame)

        # === IMPORTED DATA SECTION (Collapsible) ===
        self.section5_import_section = CollapsibleSection("Imported Data", start_collapsed=True) if CollapsibleSection else None

        self.section5_import_container = QWidget()
        self.section5_import_layout = QVBoxLayout(self.section5_import_container)
        self.section5_import_layout.setContentsMargins(8, 8, 8, 8)
        self.section5_import_layout.setSpacing(4)

        self.section5_import_placeholder = QLabel("No imported data.")
        self.section5_import_placeholder.setStyleSheet("color: #64748b; font-style: italic; font-size: 15px;")
        self.section5_import_layout.addWidget(self.section5_import_placeholder)

        self.section5_imported_entries = []

        if self.section5_import_section:
            self.section5_import_section.set_content(self.section5_import_container)
            self.section5_import_section.set_content_height(150)
            right_layout.addWidget(self.section5_import_section)
        else:
            right_layout.addWidget(self.section5_import_container)

        right_layout.addStretch()
        split.addWidget(right_container, 3)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def _update_ot_manageable_visibility(self):
        """Show/hide OT manageable radios based on concerns level."""
        concerns_val = self.ot_concerns_slider.value()
        self.ot_manageable_container.setVisible(concerns_val > 0)

    def _update_psych_manageable_visibility(self):
        """Show/hide Psychology manageable radios based on concerns level."""
        concerns_val = self.psych_concerns_slider.value()
        self.psych_manageable_container.setVisible(concerns_val > 0)

    def _update_addressing_text(self):
        """Update addressing issues text based on selections."""
        p = self._get_pronouns()
        parts = []

        # Index offence work - new format: "The work into the index offence is currently ongoing with some effectiveness."
        idx_level = self.index_offence_slider.value()
        idx_eff = self.index_offence_effectiveness.value()
        idx_progress_labels = ["not started", "started", "currently ongoing", "at an advanced stage", "completed"]
        idx_eff_labels = ["no", "some", "moderate", "significant", "high"]

        if idx_level > 0:
            text = f"The work into the index offence is {idx_progress_labels[idx_level]}"
            if idx_eff > 0:
                text += f" with {idx_eff_labels[idx_eff]} effectiveness"
            if self.index_offence_details.text().strip():
                text += f" ({self.index_offence_details.text().strip()})"
            parts.append(text + ".")

        # Risks work - new format: "Work on risk factors is currently ongoing with moderate effectiveness."
        risk_level = self.risks_work_slider.value()
        risk_eff = self.risks_work_effectiveness.value()
        risk_progress_labels = ["not started", "started", "currently ongoing", "at an advanced stage", "completed"]
        risk_eff_labels = ["no", "some", "moderate", "significant", "high"]

        if risk_level > 0:
            text = f"Work on risk factors is {risk_progress_labels[risk_level]}"
            if risk_eff > 0:
                text += f" with {risk_eff_labels[risk_eff]} effectiveness"
            if self.risks_work_details.text().strip():
                text += f" ({self.risks_work_details.text().strip()})"
            parts.append(text + ".")

        # OT activities with concerns
        if self.ot_cb.isChecked():
            ot_items = [cb.text() for cb in self.ot_checkboxes.values() if cb.isChecked()]
            if ot_items:
                ot_text = f"{p['subj']} {p['engages']} in prosocial OT activities including {', '.join(ot_items).lower()}."
                parts.append(ot_text)

                # Add concerns if any
                ot_concerns = self.ot_concerns_slider.value()
                if ot_concerns > 0:
                    concern_labels = ["no", "only minor", "some", "moderate", "significant"]
                    concern_text = f"There have been {concern_labels[ot_concerns]} concerns in these OT groups"
                    if self.ot_manageable.isChecked():
                        concern_text += " and are manageable."
                    elif self.ot_not_manageable.isChecked():
                        concern_text += " and are not manageable."
                    else:
                        concern_text += "."
                    parts.append(concern_text)

        # Psychology groups with concerns
        if self.psych_groups_cb.isChecked():
            psych_items = [cb.text() for cb in self.psych_checkboxes.values() if cb.isChecked()]
            if psych_items:
                psych_text = f"{p['subj']} {p['attends']} psychology groups including {', '.join(psych_items).lower()}."
                parts.append(psych_text)

                # Add concerns if any
                psych_concerns = self.psych_concerns_slider.value()
                if psych_concerns > 0:
                    concern_labels = ["no", "only minor", "some", "moderate", "significant"]
                    concern_text = f"There have been {concern_labels[psych_concerns]} concerns in these psychology groups"
                    if self.psych_manageable.isChecked():
                        concern_text += " and are manageable."
                    elif self.psych_not_manageable.isChecked():
                        concern_text += " and are not manageable."
                    else:
                        concern_text += "."
                    parts.append(concern_text)

        # Risk understanding
        risk_parts = []
        for key, widgets in self.risk_understanding.items():
            if widgets["checkbox"].isChecked():
                und = widgets["understanding"].currentText()
                eng_val = widgets["engagement_slider"].value()
                eng_labels = ["no", "started", "ongoing", "advanced", "completed"]
                if und != "Select...":
                    risk_parts.append(f"{widgets['checkbox'].text().lower()} ({und.lower()} understanding, {eng_labels[eng_val]} engagement)")
        if risk_parts:
            parts.append(f"Regarding risk factors: {'; '.join(risk_parts)}.")

        # Effectiveness
        for key, rb in self.effectiveness_radios.items():
            if rb.isChecked():
                parts.append(f"Overall effectiveness of interventions has been {rb.text().lower()}.")
                break

        # Relapse prevention
        relapse_val = self.relapse_slider.value()
        relapse_labels = ["nil", "started", "ongoing", "almost completed", "completed"]
        if self.relapse_refused.isChecked():
            parts.append(f"{p['subj']} {p['has']} refused relapse prevention work.")
        elif relapse_val > 0:
            parts.append(f"Relapse prevention work is {relapse_labels[relapse_val]}.")

        if parts:
            new_generated = " ".join(parts)
            self._update_text_preserving_additions(self.addressing_issues_text, new_generated, "addressing")

    def populate_section5_import_data(self, entries: list):
        """Populate the section 5 imported data panel with entries that have checkboxes."""
        # Guard check for optional layout
        if not hasattr(self, 'section5_import_layout') or not self.section5_import_layout:
            print("[MOJ-ASR] Section 5: section5_import_layout not available, skipping")
            return
        # Clear existing entries
        while self.section5_import_layout.count():
            item = self.section5_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        self.section5_imported_entries = []

        if not entries:
            self.section5_import_placeholder = QLabel("No imported data.")
            self.section5_import_placeholder.setStyleSheet("color: #64748b; font-style: italic; font-size: 15px;")
            self.section5_import_layout.addWidget(self.section5_import_placeholder)
            return

        for entry in entries:
            text = entry.get("text", "") or entry.get("content", "")
            date = entry.get("date", "") or entry.get("datetime", "")
            if not text:
                continue
            # Remove blank lines
            text = '\n'.join(line for line in text.split('\n') if line.strip())

            entry_frame = QFrame()
            entry_frame.setStyleSheet("QFrame { background: white; border: none; border-radius: 4px; }")
            entry_layout = QHBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 4, 6, 4)
            entry_layout.setSpacing(6)

            cb = QCheckBox()
            cb.setStyleSheet("QCheckBox { margin-right: 2px; }")
            entry_layout.addWidget(cb)

            content_layout = QVBoxLayout()
            content_layout.setSpacing(1)

            if date:
                date_lbl = QLabel(str(date))
                date_lbl.setStyleSheet("font-size: 22px; font-weight: 600; color: #64748b;")
                content_layout.addWidget(date_lbl)

            # Truncate text for display
            display_text = text[:150] + "..." if len(text) > 150 else text
            text_lbl = QLabel(display_text)
            text_lbl.setWordWrap(True)
            text_lbl.setStyleSheet("font-size: 16px; color: #374151;")
            content_layout.addWidget(text_lbl)

            entry_layout.addLayout(content_layout, 1)
            self.section5_import_layout.addWidget(entry_frame)

            # Store checkbox with full text
            entry_data = {"checkbox": cb, "text": text, "date": date}
            self.section5_imported_entries.append(entry_data)

            # Connect checkbox to update text
            cb.toggled.connect(self._on_section5_import_toggled)

        # Expand the import section if it exists and has CollapsibleSection
        if self.section5_import_section and hasattr(self.section5_import_section, '_toggle_collapse'):
            if self.section5_import_section.is_collapsed():
                self.section5_import_section._toggle_collapse()

    def _on_section5_import_toggled(self):
        """Handle toggling of imported section 5 entries - add/remove from text area."""
        current_text = self.addressing_issues_text.toPlainText()

        # Collect all checked entries
        checked_texts = []
        for entry in self.section5_imported_entries:
            if entry["checkbox"].isChecked():
                date = entry.get("date", "")
                text = entry["text"]
                if date:
                    checked_texts.append(f"[{date}] {text}")
                else:
                    checked_texts.append(text)

        # Find the import section marker or append
        import_marker = "\n\n--- Imported Notes ---\n"
        if import_marker in current_text:
            # Replace the imported section
            base_text = current_text.split(import_marker)[0]
        else:
            base_text = current_text

        if checked_texts:
            new_text = base_text.rstrip() + import_marker + "\n".join(checked_texts)
        else:
            new_text = base_text.rstrip()

        self.addressing_issues_text.setPlainText(new_text)

    # ================================================================
    # SECTION 6: Patient Attitude (Last 12 Months)
    # ================================================================
    def _build_section_6_patient_attitude(self):
        frame = self._create_section_frame("Patient Attitude (Last 12 Months)", "6", "#7c3aed")
        layout = frame.layout()

        prompt_lbl = QLabel("Describe the patient's attitude in the last 12 months:")
        prompt_lbl.setWordWrap(True)
        prompt_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(prompt_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        # Left side: text area
        self.patient_attitude_text = self._create_text_edit("", 400)
        self.patient_attitude_text.setMinimumWidth(350)
        split.addWidget(self.patient_attitude_text, 2)

        # Right side: scrollable controls panel
        right_container = QWidget()
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.setSpacing(8)

        # === COMPLIANCE GRID SECTION (like Section 15 from tribunal) ===
        self.compliance_section = CollapsibleSection("Understanding & Compliance") if CollapsibleSection else None

        compliance_content = QWidget()
        compliance_content.setObjectName("compliance_content")
        compliance_content.setStyleSheet("""
            QWidget#compliance_content {
                background: rgba(255,255,255,0.95);
                border-radius: 8px;
            }
        """)
        compliance_layout = QVBoxLayout(compliance_content)
        compliance_layout.setContentsMargins(12, 12, 12, 12)
        compliance_layout.setSpacing(12)

        # Grid for treatments (matching tribunal style)
        grid = QGridLayout()
        grid.setSpacing(8)

        # Headers
        header_treatment = QLabel("Treatment")
        header_treatment.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        header_understanding = QLabel("Understanding")
        header_understanding.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        header_compliance = QLabel("Compliance")
        header_compliance.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")

        grid.addWidget(header_treatment, 0, 0)
        grid.addWidget(header_understanding, 0, 1)
        grid.addWidget(header_compliance, 0, 2)

        # Treatment rows
        UNDERSTANDING_OPTIONS = ["Select...", "good", "fair", "poor"]
        COMPLIANCE_OPTIONS = ["Select...", "full", "reasonable", "partial", "nil"]

        self.attitude_treatments = {}
        treatment_names = ["Medical", "Nursing", "Psychology", "OT", "Social Work"]

        for i, name in enumerate(treatment_names, 1):
            key = name.lower().replace(" ", "_")

            # Label
            lbl = QLabel(name)
            lbl.setStyleSheet("font-size: 16px; color: #374151;")
            grid.addWidget(lbl, i, 0)

            # Understanding dropdown
            understanding = QComboBox()
            understanding.addItems(UNDERSTANDING_OPTIONS)
            understanding.currentIndexChanged.connect(self._update_attitude_text)
            grid.addWidget(understanding, i, 1)

            # Compliance dropdown
            compliance = QComboBox()
            compliance.addItems(COMPLIANCE_OPTIONS)
            compliance.currentIndexChanged.connect(self._update_attitude_text)
            grid.addWidget(compliance, i, 2)

            self.attitude_treatments[key] = {
                "understanding": understanding,
                "compliance": compliance
            }

        compliance_layout.addLayout(grid)
        compliance_layout.addStretch()

        if self.compliance_section:
            self.compliance_section.set_content(compliance_content)
            self.compliance_section.set_content_height(220)
            right_layout.addWidget(self.compliance_section)
        else:
            right_layout.addWidget(compliance_content)

        # === IMPORTED DATA SECTION (FixedDataPanel style from Section 14) ===
        self.section6_import_section = CollapsibleSection("Imported Data", start_collapsed=True) if CollapsibleSection else None

        import_panel = QWidget()
        import_panel_layout = QVBoxLayout(import_panel)
        import_panel_layout.setContentsMargins(8, 8, 8, 8)
        import_panel_layout.setSpacing(8)

        # Subtitle (fixed height - doesn't stretch)
        self.section6_subtitle = QLabel("Progress and mental state data from notes")
        self.section6_subtitle.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        self.section6_subtitle.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        import_panel_layout.addWidget(self.section6_subtitle)

        # Scrollable content area (this stretches)
        section6_scroll = QScrollArea()
        section6_scroll.setWidgetResizable(True)
        section6_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                border-radius: 8px;
                background: white;
            }
        """)
        section6_scroll.setMinimumHeight(150)
        section6_scroll.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        # Content widget
        self.section6_content_widget = QWidget()
        self.section6_import_layout = QVBoxLayout(self.section6_content_widget)
        self.section6_import_layout.setContentsMargins(12, 12, 12, 12)
        self.section6_import_layout.setSpacing(8)
        section6_scroll.setWidget(self.section6_content_widget)

        # Placeholder
        self.section6_import_placeholder = QLabel("No imported data. Use Import File to upload data.")
        self.section6_import_placeholder.setStyleSheet("color: #9ca3af; font-style: italic;")
        self.section6_import_layout.addWidget(self.section6_import_placeholder)

        import_panel_layout.addWidget(section6_scroll, 1)

        # Button row
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(8)

        # Clear button
        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #f3f4f6;
                color: #374151;
                border: 1px solid #d1d5db;
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 500;
                font-size: 15px;
            }
            QPushButton:hover { background: #e5e7eb; }
        """)
        clear_btn.clicked.connect(self._clear_section6_imports)
        btn_layout.addWidget(clear_btn)

        # Send to Report button (sends selected/checked entries)
        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #8b5cf6;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 15px;
            }
            QPushButton:hover { background: #7c3aed; }
        """)
        send_btn.clicked.connect(self._send_section6_to_report)
        btn_layout.addWidget(send_btn)

        import_panel_layout.addLayout(btn_layout)

        self.section6_imported_entries = []
        self.section6_checkboxes = []

        if self.section6_import_section:
            self.section6_import_section.set_content(import_panel)
            self.section6_import_section.set_content_height(350)
            right_layout.addWidget(self.section6_import_section)
        else:
            right_layout.addWidget(import_panel)

        right_layout.addStretch()
        split.addWidget(right_container, 3)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

        # Apply any extracted patient details (compliance/understanding) now that widgets exist
        self._apply_extracted_patient_details()

    def _update_attitude_text(self):
        """Update patient attitude text based on compliance grid selections."""
        p = self._get_pronouns()
        parts = []

        # Medical
        med = self.attitude_treatments["medical"]
        med_u = med["understanding"].currentText()
        med_c = med["compliance"].currentText()
        if med_u != "Select..." and med_c != "Select...":
            u_phrase = self._attitude_understanding_phrase(med_u, "medical", p)
            c_phrase = self._attitude_compliance_phrase(med_c, p)
            if u_phrase and c_phrase:
                parts.append(f"{u_phrase} {c_phrase}.")

        # Nursing
        nursing = self.attitude_treatments["nursing"]
        nursing_u = nursing["understanding"].currentText()
        nursing_c = nursing["compliance"].currentText()
        if nursing_u != "Select..." and nursing_c != "Select...":
            phrase = self._nursing_attitude_phrase(nursing_u, nursing_c, p)
            if phrase:
                parts.append(phrase)

        # Psychology
        psych = self.attitude_treatments["psychology"]
        psych_u = psych["understanding"].currentText()
        psych_c = psych["compliance"].currentText()
        if psych_u != "Select..." and psych_c != "Select...":
            phrase = self._psychology_attitude_phrase(psych_u, psych_c, p)
            if phrase:
                parts.append(phrase)

        # OT
        ot = self.attitude_treatments["ot"]
        ot_u = ot["understanding"].currentText()
        ot_c = ot["compliance"].currentText()
        if ot_u != "Select..." and ot_c != "Select...":
            phrase = self._ot_attitude_phrase(ot_u, ot_c, p)
            if phrase:
                parts.append(phrase)

        # Social Work
        sw = self.attitude_treatments["social_work"]
        sw_u = sw["understanding"].currentText()
        sw_c = sw["compliance"].currentText()
        if sw_u != "Select..." and sw_c != "Select...":
            phrase = self._social_work_attitude_phrase(sw_u, sw_c, p)
            if phrase:
                parts.append(phrase)

        if parts:
            new_generated = " ".join(parts)
            self._update_text_preserving_additions(self.patient_attitude_text, new_generated, "attitude")

    def _attitude_understanding_phrase(self, level: str, treatment: str, p: dict) -> str:
        """Generate understanding phrase based on level."""
        if level == "good":
            return f"{p['subj']} {p['has']} good understanding of {p['pos_l']} {treatment} treatment"
        elif level == "fair":
            return f"{p['subj']} {p['has']} some understanding of {p['pos_l']} {treatment} treatment"
        elif level == "poor":
            return f"{p['subj']} {p['has']} limited understanding of {p['pos_l']} {treatment} treatment"
        return ""

    def _attitude_compliance_phrase(self, level: str, p: dict) -> str:
        """Generate compliance phrase based on level."""
        if level == "full":
            return "and compliance is full"
        elif level == "reasonable":
            return "and compliance is reasonable"
        elif level == "partial":
            return "but compliance is partial"
        elif level == "nil":
            return "and compliance is nil"
        return ""

    def _nursing_attitude_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural nursing phrase."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['engages']} well with nursing staff and {p['has']} insight into the need for nursing input."
        elif understanding == "good" and compliance == "partial":
            return f"{p['subj']} understands the role of nursing but engagement is inconsistent."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['has']} some understanding of nursing care and {p['engages']} reasonably well."
        elif understanding == "fair" and compliance == "partial":
            return f"{p['subj']} {p['has']} some understanding of nursing input but {p['engages']} only partially."
        elif understanding == "poor" or compliance == "nil":
            return f"{p['subj']} {p['has']} limited insight into the need for nursing care and {p['is']} not engaging meaningfully."
        return ""

    def _psychology_attitude_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural psychology phrase."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['engages']} in psychology sessions and sees the benefit of this work."
        elif understanding == "good" and compliance == "partial":
            return f"{p['subj']} understands the purpose of psychology but compliance with sessions is limited."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['has']} some understanding of psychology and attends sessions regularly."
        elif understanding == "fair" and compliance == "partial":
            return f"{p['subj']} also {p['engages']} in psychology sessions but the compliance with these is limited."
        elif understanding == "poor" or compliance == "nil":
            return f"{p['subj']} {p['has']} limited insight into the need for psychology and {p['is']} not engaging with sessions."
        return ""

    def _ot_attitude_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural OT phrase."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"With respect to OT, {p['subj_l']} {p['engages']} well and sees the benefit of activities."
        elif understanding == "good" and compliance == "partial":
            return f"With respect to OT, {p['subj_l']} understands the purpose but engagement is inconsistent."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"With respect to OT, {p['subj_l']} {p['has']} some understanding and participates in activities."
        elif understanding == "fair" and compliance == "partial":
            return f"With respect to OT, {p['subj_l']} {p['has']} some insight but engagement is limited."
        elif understanding == "poor" or compliance == "nil":
            return f"With respect to OT, {p['subj_l']} {p['is']} not engaging and doesn't see the need to."
        return ""

    def _social_work_attitude_phrase(self, understanding: str, compliance: str, p: dict) -> str:
        """Generate natural social work phrase."""
        if understanding == "good" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['engages']} well with the social worker and understands {p['pos_l']} social circumstances."
        elif understanding == "good" and compliance == "partial":
            return f"{p['subj']} understands the social worker's role but engagement is inconsistent."
        elif understanding == "fair" and compliance in ("full", "reasonable"):
            return f"{p['subj']} {p['has']} some understanding of social work input and {p['engages']} when available."
        elif understanding == "fair" and compliance == "partial":
            return f"{p['subj']} occasionally sees the social worker and {p['engages']} partially."
        elif understanding == "poor" or compliance == "nil":
            return f"{p['subj']} {p['has']} limited engagement with social work and doesn't see the relevance."
        return ""

    def _populate_section6_filtered_data(self) -> int:
        """
        Populate section 6 with filtered notes from 1 year before the last entry.
        Filters: Mental State, Compliance, Attendance, Hospital Admissions.
        Shows first 2 lines + lines containing filter keywords.
        Returns the count of entries found.
        """
        from datetime import datetime, timedelta

        raw_notes = self._extracted_raw_notes
        if not raw_notes:
            print("[MOJ-ASR] Section 6: No raw notes available")
            self.populate_section6_import_data([])
            return 0

        # Parse date helper
        def parse_note_date(date_val):
            if isinstance(date_val, datetime):
                return date_val
            if not date_val:
                return None
            date_str = str(date_val).strip()
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d", "%d/%m/%Y %H:%M", "%Y-%m-%d %H:%M:%S"]:
                try:
                    return datetime.strptime(date_str.split()[0] if ' ' in date_str else date_str, fmt.split()[0])
                except:
                    pass
            return None

        # Find the most recent note date
        all_dates = []
        for n in raw_notes:
            dt = parse_note_date(n.get("date") or n.get("datetime"))
            if dt:
                all_dates.append(dt)

        if not all_dates:
            print("[MOJ-ASR] Section 6: No parseable dates in notes")
            self.populate_section6_import_data([])
            return 0

        most_recent = max(all_dates)
        one_year_cutoff = most_recent - timedelta(days=365)

        print(f"[MOJ-ASR] Section 6: Most recent note: {most_recent.strftime('%d/%m/%Y')}")
        print(f"[MOJ-ASR] Section 6: 1-year cutoff: {one_year_cutoff.strftime('%d/%m/%Y')}")

        # Keywords for each category (Admission is handled separately via timeline_builder)
        FILTERS = {
            "Mental State": [
                "mental state", "mse", "mood", "affect", "thought", "delusion", "hallucination",
                "psychotic", "psychosis", "depression", "depressed", "anxiety", "anxious",
                "paranoid", "paranoia", "voices", "hearing voices", "suicidal", "self-harm",
                "presentation", "eye contact", "speech", "insight", "agitated", "elated",
                "flat affect", "blunted", "congruent", "incongruent", "thought disorder"
            ],
            "Compliance": [
                "compliant", "compliance", "non-compliant", "medication", "taking medication",
                "refusing medication", "depot", "injection", "oral medication", "tablets",
                "concordant", "concordance", "adherent", "adherence", "treatment plan",
                "accepted", "declined", "refused", "non-concordant", "not taking"
            ],
            "Attendance": [
                "appointment", "attended", "did not attend", "dna", "failed to attend",
                "engagement", "engaged", "engaging", "not engaging", "ward round",
                "review", "session", "meeting", "psychology", "ot session", "group",
                "1:1", "one to one", "1-1", "cpa", "absent", "present", "arrived"
            ]
        }

        def get_matching_categories(text: str) -> list:
            """Return list of categories that match the text."""
            if not text:
                return []
            text_lower = text.lower()
            matched = []
            for cat, keywords in FILTERS.items():
                if any(kw in text_lower for kw in keywords):
                    matched.append(cat)
            return matched

        def extract_relevant_lines(text: str, categories: list) -> str:
            """Extract first 2 lines + lines containing filter keywords."""
            if not text:
                return ""

            lines = text.split('\n')
            lines = [l.strip() for l in lines if l.strip()]  # Remove empty lines

            if not lines:
                return text[:500]

            # Get first 2 lines
            result_lines = []
            first_two = lines[:2]
            result_lines.extend(first_two)

            # Get all keywords for matched categories
            all_keywords = []
            for cat in categories:
                all_keywords.extend(FILTERS.get(cat, []))

            # Find lines containing keywords (after first 2)
            for i, line in enumerate(lines[2:], start=2):
                line_lower = line.lower()
                if any(kw in line_lower for kw in all_keywords):
                    # Add separator if there's a gap
                    if result_lines and i > len(result_lines):
                        if result_lines[-1] != "...":
                            result_lines.append("...")
                    result_lines.append(line)

            return '\n'.join(result_lines)

        # Filter and process notes
        section6_entries = []
        seen_texts = set()

        for i, n in enumerate(raw_notes):
            dt = parse_note_date(n.get("date") or n.get("datetime"))
            if not dt or dt < one_year_cutoff:
                continue

            full_text = (n.get("content") or n.get("text") or n.get("preview") or "").strip()
            if not full_text or full_text in seen_texts:
                continue

            # Check if note matches any filter
            categories = get_matching_categories(full_text)
            if not categories:
                continue

            seen_texts.add(full_text)

            # Extract relevant lines
            display_text = extract_relevant_lines(full_text, categories)

            section6_entries.append({
                "text": display_text,
                "full_text": full_text,  # Keep full text for sending to report
                "date": str(n.get("date") or n.get("datetime", "")),
                "categories": categories,
                "sort_date": dt
            })

        # Detect hospital admissions using timeline_builder (not keywords)
        try:
            from timeline_builder import build_timeline
            import traceback

            print(f"[MOJ-ASR] Section 6: Building timeline from {len(raw_notes)} raw notes...")

            # Build timeline from ALL notes (not just 1-year) to detect all admissions
            timeline_notes = []
            for n in raw_notes:
                dt = parse_note_date(n.get("date") or n.get("datetime"))
                if dt:
                    text = (n.get("content") or n.get("text") or n.get("preview") or "").strip()
                    timeline_notes.append({
                        "date": dt,
                        "content": text,
                        "text": text,
                        "preview": n.get("preview", ""),
                        "type": n.get("type", ""),
                        "source": n.get("source", "")
                    })

            print(f"[MOJ-ASR] Section 6: Prepared {len(timeline_notes)} notes for timeline")

            if timeline_notes:
                episodes = build_timeline(timeline_notes)
                print(f"[MOJ-ASR] Section 6: Timeline returned {len(episodes)} episodes")

                # Store episodes for Section 4 to access
                self._timeline_episodes = episodes

                for ep in episodes:
                    print(f"[MOJ-ASR] Section 6:   Episode: {ep.get('type')} from {ep.get('start')} to {ep.get('end')}")

                inpatient_episodes = [e for e in episodes if e.get("type") == "inpatient"]
                print(f"[MOJ-ASR] Section 6: Found {len(inpatient_episodes)} inpatient episodes")

                # Build list of relevant admission periods (those overlapping with 1-year window)
                relevant_admissions = []
                if inpatient_episodes:
                    for ep in inpatient_episodes:
                        ep_start = ep["start"]
                        ep_end = ep["end"]

                        print(f"[MOJ-ASR] Section 6: Processing admission: {ep_start} to {ep_end}")

                        # Convert date objects to datetime if needed for comparison
                        ep_start_dt = ep_start
                        ep_end_dt = ep_end
                        if not isinstance(ep_start_dt, datetime):
                            from datetime import date as date_type
                            if isinstance(ep_start_dt, date_type):
                                ep_start_dt = datetime.combine(ep_start_dt, datetime.min.time())
                        if not isinstance(ep_end_dt, datetime):
                            from datetime import date as date_type
                            if isinstance(ep_end_dt, date_type):
                                ep_end_dt = datetime.combine(ep_end_dt, datetime.min.time())

                        # Check if admission OVERLAPS with 1-year window
                        if ep_end_dt < one_year_cutoff:
                            print(f"[MOJ-ASR] Section 6:   SKIPPED - admission ends before 1-year cutoff ({one_year_cutoff.strftime('%d/%m/%Y')})")
                            continue

                        relevant_admissions.append({
                            "start": ep_start_dt,
                            "end": ep_end_dt,
                            "label": ep.get("label", "Admission")
                        })

                        start_str = ep["start"].strftime("%d/%m/%Y") if hasattr(ep["start"], "strftime") else str(ep["start"])
                        end_str = ep["end"].strftime("%d/%m/%Y") if hasattr(ep["end"], "strftime") else str(ep["end"])
                        label = ep.get("label", "Admission")

                        # Add admission summary entry
                        admission_text = f"Hospital admission from {start_str} to {end_str}"
                        section6_entries.append({
                            "text": f"[{label}] {admission_text}",
                            "full_text": admission_text,
                            "date": start_str,
                            "categories": ["Admission"],
                            "sort_date": ep_start_dt
                        })
                        print(f"[MOJ-ASR] Section 6:   ADDED admission: {start_str} to {end_str}")

                # Tag existing entries with "Admission" if they fall within an admission period
                if relevant_admissions:
                    tagged_count = 0
                    for entry in section6_entries:
                        entry_date = entry.get("sort_date")
                        if not entry_date:
                            continue

                        # Check if entry date falls within any admission period
                        for adm in relevant_admissions:
                            if adm["start"] <= entry_date <= adm["end"]:
                                if "Admission" not in entry.get("categories", []):
                                    entry["categories"] = entry.get("categories", []) + ["Admission"]
                                    tagged_count += 1
                                break

                    print(f"[MOJ-ASR] Section 6: Tagged {tagged_count} entries with 'Admission' category")
                else:
                    print(f"[MOJ-ASR] Section 6: No inpatient episodes found in timeline")
            else:
                print(f"[MOJ-ASR] Section 6: No notes prepared for timeline")

        except ImportError as e:
            print(f"[MOJ-ASR] Section 6: Could not import timeline_builder: {e}")
        except Exception as e:
            import traceback
            print(f"[MOJ-ASR] Section 6: Timeline error: {e}")
            traceback.print_exc()

        # Sort by date (most recent first)
        section6_entries.sort(key=lambda x: x.get("sort_date", datetime.min), reverse=True)

        # Remove sort_date before passing
        for e in section6_entries:
            e.pop("sort_date", None)

        print(f"[MOJ-ASR] Section 6: Found {len(section6_entries)} filtered entries in 1-year period")
        self.populate_section6_import_data(section6_entries)
        return len(section6_entries)

    def populate_section6_import_data(self, entries: list):
        """Populate the section 6 imported data panel with entries and checkboxes."""
        # Store entries
        self.section6_imported_entries = entries
        self.section6_checkboxes = []  # Store checkbox references

        # Check if UI elements exist (old collapsible section UI may not be built)
        if not hasattr(self, 'section6_import_layout') or not self.section6_import_layout:
            print("[MOJ-ASR] Section 6 import UI not available - skipping populate")
            return

        # Update subtitle (if UI element exists)
        if hasattr(self, 'section6_subtitle') and self.section6_subtitle:
            if entries:
                self.section6_subtitle.setText(f"Showing {len(entries)} entries from last 1 year - tick to include in letter")
            else:
                self.section6_subtitle.setText("No entries found")

        # Clear existing content
        while self.section6_import_layout.count():
            item = self.section6_import_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not entries:
            placeholder = QLabel("No imported data. Use Import File to upload data.")
            placeholder.setStyleSheet("color: #9ca3af; font-style: italic;")
            self.section6_import_layout.addWidget(placeholder)
            return

        # Select All / Deselect All buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        select_all_btn = QPushButton("Select All")
        select_all_btn.setStyleSheet("""
            QPushButton {
                background: #3b82f6;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 6px 12px;
                font-size: 15px;
                font-weight: 600;
            }
            QPushButton:hover { background: #2563eb; }
        """)
        select_all_btn.clicked.connect(self._select_all_section6)
        btn_row.addWidget(select_all_btn)

        deselect_all_btn = QPushButton("Deselect All")
        deselect_all_btn.setStyleSheet("""
            QPushButton {
                background: #6b7280;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 6px 12px;
                font-size: 15px;
                font-weight: 600;
            }
            QPushButton:hover { background: #4b5563; }
        """)
        deselect_all_btn.clicked.connect(self._deselect_all_section6)
        btn_row.addWidget(deselect_all_btn)

        btn_row.addStretch()
        self.section6_import_layout.addLayout(btn_row)

        # Category tag colors
        CATEGORY_COLORS = {
            "Mental State": "#3b82f6",   # Blue
            "Compliance": "#10b981",      # Green
            "Attendance": "#f59e0b",      # Amber
            "Admission": "#ef4444"        # Red
        }

        # Add each entry with checkbox
        for i, entry in enumerate(entries):
            text = entry.get("text", "") or entry.get("content", "")
            date = entry.get("date", "") or entry.get("datetime", "")
            categories = entry.get("categories", [])
            if not text:
                continue

            entry_frame = QFrame()
            entry_frame.setStyleSheet("""
                QFrame {
                    background: #f9fafb;
                    border: none;
                    border-radius: 6px;
                }
            """)
            entry_layout = QHBoxLayout(entry_frame)
            entry_layout.setContentsMargins(8, 8, 8, 8)
            entry_layout.setSpacing(8)

            # Checkbox
            checkbox = QCheckBox()
            checkbox.setStyleSheet("QCheckBox::indicator { width: 18px; height: 18px; }")
            entry_layout.addWidget(checkbox)
            self.section6_checkboxes.append((checkbox, entry))

            # Content area
            content_layout = QVBoxLayout()
            content_layout.setSpacing(4)

            # Header row with date and category tags
            header_row = QHBoxLayout()
            header_row.setSpacing(6)

            # Date header if available
            if date:
                date_label = QLabel(str(date))
                date_label.setStyleSheet("font-size: 15px; font-weight: 600; color: #6b7280;")
                header_row.addWidget(date_label)

            # Category tags
            for cat in categories:
                color = CATEGORY_COLORS.get(cat, "#6b7280")
                tag = QLabel(cat)
                tag.setStyleSheet(f"""
                    background: {color};
                    color: white;
                    font-size: 22px;
                    font-weight: 600;
                    padding: 2px 6px;
                    border-radius: 3px;
                """)
                header_row.addWidget(tag)

            header_row.addStretch()
            content_layout.addLayout(header_row)

            # Content text
            text_label = QLabel(text)
            text_label.setWordWrap(True)
            text_label.setStyleSheet("font-size: 16px; color: #374151;")
            content_layout.addWidget(text_label)

            entry_layout.addLayout(content_layout, 1)
            self.section6_import_layout.addWidget(entry_frame)

        # Add stretch at end
        self.section6_import_layout.addStretch()

        # Expand the import section if it exists
        if self.section6_import_section and hasattr(self.section6_import_section, '_toggle_collapse'):
            if self.section6_import_section.is_collapsed():
                self.section6_import_section._toggle_collapse()

    def _select_all_section6(self):
        """Select all checkboxes in section 6."""
        for checkbox, _ in self.section6_checkboxes:
            checkbox.setChecked(True)

    def _deselect_all_section6(self):
        """Deselect all checkboxes in section 6."""
        for checkbox, _ in self.section6_checkboxes:
            checkbox.setChecked(False)

    def _clear_section6_imports(self):
        """Clear the section 6 imported data."""
        self.populate_section6_import_data([])
        self.section6_checkboxes = []

    def _send_section6_to_report(self):
        """Send selected (checked) section 6 entries to the text area."""
        if not hasattr(self, 'section6_checkboxes') or not self.section6_checkboxes:
            return

        # Combine only checked entry texts (use full_text if available for complete note)
        texts = []
        for checkbox, entry in self.section6_checkboxes:
            if checkbox.isChecked():
                # Use full_text (complete note) if available, otherwise use display text
                text = (entry.get("full_text", "") or entry.get("text", "") or entry.get("content", "")).strip()
                date_str = entry.get("date", "")
                if text:
                    if date_str:
                        texts.append(f"[{date_str}] {text}")
                    else:
                        texts.append(text)

        if not texts:
            QMessageBox.information(self, "No Selection", "Please tick at least one entry to send to the letter.")
            return

        combined = "\n\n".join(texts)

        # Append to existing text
        current = self.patient_attitude_text.toPlainText()
        if current:
            self.patient_attitude_text.setPlainText(current + "\n\n--- Imported Notes ---\n" + combined)
        else:
            self.patient_attitude_text.setPlainText(combined)

        QMessageBox.information(self, "Sent", f"{len(texts)} entries sent to letter.")

    # ================================================================
    # SECTION 7: Capacity Issues
    # ================================================================
    def _build_section_7_capacity(self):
        frame = self._create_section_frame("Capacity Issues", "7", "#dc2626")
        layout = frame.layout()

        split = QHBoxLayout()
        split.setSpacing(16)

        # Left side - text area
        left_layout = QVBoxLayout()
        self.capacity_text = self._create_text_edit("", 140)
        left_layout.addWidget(self.capacity_text)
        split.addLayout(left_layout, 3)

        # Right side - domain dropdown, capacity radio buttons, and actions
        right_layout = QVBoxLayout()
        right_layout.setSpacing(8)

        from PySide6.QtWidgets import QButtonGroup, QRadioButton, QComboBox, QStackedWidget

        # Store capacity entries per domain: {"medication": "has", "finances": "lacks", ...}
        self.capacity_entries = {}
        # Store actions per domain
        self.capacity_domain_actions = {}

        # Domain dropdown FIRST (primary entry)
        capacity_label = QLabel("Capacity:")
        capacity_label.setStyleSheet("font-weight: 600; color: #dc2626; font-size: 16px;")
        right_layout.addWidget(capacity_label)

        domain_label = QLabel("Domain:")
        domain_label.setStyleSheet("font-size: 15px; color: #6b7280;")
        right_layout.addWidget(domain_label)

        self.capacity_domain_combo = QComboBox()
        self.capacity_domain_combo.addItems([
            "medication",
            "finances",
            "residence",
            "self-care"
        ])
        self.capacity_domain_combo.setStyleSheet("font-size: 15px;")
        right_layout.addWidget(self.capacity_domain_combo)

        # Radio buttons for has/lacks capacity
        self.capacity_group = QButtonGroup(self)
        self.capacity_group.setExclusive(True)

        self.capacity_has_rb = QRadioButton("Has capacity")
        self.capacity_lacks_rb = QRadioButton("Lacks capacity")

        self.capacity_has_rb.setStyleSheet("""
            QRadioButton {
                font-size: 15px;
                color: #374151;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 16px;
                height: 16px;
            }
        """)
        self.capacity_lacks_rb.setStyleSheet("""
            QRadioButton {
                font-size: 15px;
                color: #374151;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 16px;
                height: 16px;
            }
        """)

        self.capacity_group.addButton(self.capacity_has_rb)
        self.capacity_group.addButton(self.capacity_lacks_rb)

        right_layout.addWidget(self.capacity_has_rb)
        right_layout.addWidget(self.capacity_lacks_rb)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.Shape.HLine)
        sep.setStyleSheet("background: #e5e7eb;")
        sep.setFixedHeight(1)
        right_layout.addWidget(sep)

        # Action taken label
        action_label = QLabel("Action taken:")
        action_label.setStyleSheet("font-weight: 600; color: #dc2626; font-size: 16px;")
        right_layout.addWidget(action_label)

        # Stacked widget for domain-specific actions
        self.action_stack = QStackedWidget()

        # Page 0: Medication - SOAD requested Yes/No
        med_page = QWidget()
        med_layout = QVBoxLayout(med_page)
        med_layout.setContentsMargins(0, 0, 0, 0)
        med_layout.setSpacing(4)

        self.soad_group = QButtonGroup(self)
        self.soad_group.setExclusive(True)
        self.soad_yes_rb = QRadioButton("SOAD requested - Yes")
        self.soad_no_rb = QRadioButton("SOAD requested - No")
        self.soad_yes_rb.setStyleSheet("""
            QRadioButton {
                font-size: 15px;
                color: #374151;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 16px;
                height: 16px;
            }
        """)
        self.soad_no_rb.setStyleSheet("""
            QRadioButton {
                font-size: 15px;
                color: #374151;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 16px;
                height: 16px;
            }
        """)
        self.soad_group.addButton(self.soad_yes_rb)
        self.soad_group.addButton(self.soad_no_rb)
        med_layout.addWidget(self.soad_yes_rb)
        med_layout.addWidget(self.soad_no_rb)
        med_layout.addStretch()
        self.action_stack.addWidget(med_page)

        # Page 1: Finances - IMCA checkbox + radio buttons for Guardianship/Appointeeship
        fin_page = QWidget()
        fin_layout = QVBoxLayout(fin_page)
        fin_layout.setContentsMargins(0, 0, 0, 0)
        fin_layout.setSpacing(4)

        # IMCA remains a checkbox (independent option)
        self.fin_imca_cb = QCheckBox("IMCA")
        self.fin_imca_cb.setStyleSheet("font-size: 15px; color: #374151;")
        self.fin_imca_cb.stateChanged.connect(self._update_capacity_text)
        fin_layout.addWidget(self.fin_imca_cb)

        # Radio button group for Guardianship/Appointeeship (mutually exclusive)
        self.fin_type_group = QButtonGroup(self)
        self.fin_none_rb = QRadioButton("None")
        self.fin_guardianship_rb = QRadioButton("Guardianship")
        self.fin_appointeeship_rb = QRadioButton("Appointeeship")
        self.fin_informal_rb = QRadioButton("Informal Appointeeship")
        self.fin_none_rb.setChecked(True)  # Default to None
        for rb in [self.fin_none_rb, self.fin_guardianship_rb, self.fin_appointeeship_rb, self.fin_informal_rb]:
            rb.setStyleSheet("font-size: 15px; color: #374151;")
            rb.toggled.connect(self._update_capacity_text)
            self.fin_type_group.addButton(rb)
            fin_layout.addWidget(rb)
        fin_layout.addStretch()
        self.action_stack.addWidget(fin_page)

        # Page 2: Residence - Best Interest Meeting, IMCA, DOLS, COP
        res_page = QWidget()
        res_layout = QVBoxLayout(res_page)
        res_layout.setContentsMargins(0, 0, 0, 0)
        res_layout.setSpacing(4)

        self.res_best_interest_cb = QCheckBox("Best Interest Meeting")
        self.res_imca_cb = QCheckBox("IMCA")
        self.res_dols_cb = QCheckBox("DOLS")
        self.res_cop_cb = QCheckBox("COP")
        for cb in [self.res_best_interest_cb, self.res_imca_cb, self.res_dols_cb, self.res_cop_cb]:
            cb.setStyleSheet("font-size: 15px; color: #374151;")
            cb.stateChanged.connect(self._update_capacity_text)
            res_layout.addWidget(cb)
        res_layout.addStretch()
        self.action_stack.addWidget(res_page)

        # Page 3: Self-care - Same as residence
        self_page = QWidget()
        self_layout = QVBoxLayout(self_page)
        self_layout.setContentsMargins(0, 0, 0, 0)
        self_layout.setSpacing(4)

        self.self_best_interest_cb = QCheckBox("Best Interest Meeting")
        self.self_imca_cb = QCheckBox("IMCA")
        self.self_dols_cb = QCheckBox("DOLS")
        self.self_cop_cb = QCheckBox("COP")
        for cb in [self.self_best_interest_cb, self.self_imca_cb, self.self_dols_cb, self.self_cop_cb]:
            cb.setStyleSheet("font-size: 15px; color: #374151;")
            cb.stateChanged.connect(self._update_capacity_text)
            self_layout.addWidget(cb)
        self_layout.addStretch()
        self.action_stack.addWidget(self_page)

        right_layout.addWidget(self.action_stack)

        # Hide actions initially (only show when lacks capacity)
        action_label.setVisible(False)
        self.action_stack.setVisible(False)
        self._action_label = action_label  # Store reference

        # Connect SOAD radios to update text
        self.soad_yes_rb.toggled.connect(self._update_capacity_text)
        self.soad_no_rb.toggled.connect(self._update_capacity_text)

        def capacity_changed():
            btn = self.capacity_group.checkedButton()
            if not btn:
                return
            status = "has" if btn is self.capacity_has_rb else "lacks"
            domain = self.capacity_domain_combo.currentText()
            # Store entry for this domain
            self.capacity_entries[domain] = status

            # Show/hide actions based on capacity status
            if status == "lacks":
                self._action_label.setVisible(True)
                self.action_stack.setVisible(True)
            else:
                self._action_label.setVisible(False)
                self.action_stack.setVisible(False)

            self._update_capacity_text()

        def domain_changed(text):
            # Switch action stack page
            domain_to_page = {"medication": 0, "finances": 1, "residence": 2, "self-care": 3}
            self.action_stack.setCurrentIndex(domain_to_page.get(text, 0))

            # Restore capacity radio state for this domain
            self.capacity_group.setExclusive(False)
            self.capacity_has_rb.setChecked(False)
            self.capacity_lacks_rb.setChecked(False)
            self.capacity_group.setExclusive(True)

            # Hide actions by default when switching domains
            self._action_label.setVisible(False)
            self.action_stack.setVisible(False)

            # Restore if entry exists for this domain
            if text in self.capacity_entries:
                status = self.capacity_entries[text]
                self.capacity_group.setExclusive(False)
                if status == "has":
                    self.capacity_has_rb.setChecked(True)
                else:
                    self.capacity_lacks_rb.setChecked(True)
                    # Show actions if lacks capacity
                    self._action_label.setVisible(True)
                    self.action_stack.setVisible(True)
                self.capacity_group.setExclusive(True)

        self.capacity_has_rb.toggled.connect(capacity_changed)
        self.capacity_lacks_rb.toggled.connect(capacity_changed)
        self.capacity_domain_combo.currentTextChanged.connect(domain_changed)

        right_layout.addStretch()

        right_container = QWidget()
        right_container.setLayout(right_layout)
        right_container.setFixedWidth(320)
        split.addWidget(right_container)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def _update_capacity_text(self):
        """Update capacity text based on additive capacity entries and domain-specific actions."""
        parts = []

        # Build capacity statements from all entries
        if self.capacity_entries:
            for domain, status in self.capacity_entries.items():
                if status == "has":
                    parts.append(
                        f"Capacity assessment (understands information, retains it, "
                        f"weighs up pros and cons, and can communicate wishes) was "
                        f"carried out for {domain} and the patient is noted to "
                        f"have capacity."
                    )
                else:
                    parts.append(
                        f"Capacity assessment (understands information, retains it, "
                        f"weighs up pros and cons, and can communicate wishes) was "
                        f"carried out for {domain} and the patient is noted to "
                        f"lack capacity."
                    )

                # Add domain-specific actions
                domain_actions = []

                if domain == "medication":
                    if hasattr(self, 'soad_yes_rb') and self.soad_yes_rb.isChecked():
                        domain_actions.append("SOAD has been requested")
                    elif hasattr(self, 'soad_no_rb') and self.soad_no_rb.isChecked():
                        domain_actions.append("SOAD has not been requested")

                elif domain == "finances":
                    if hasattr(self, 'fin_imca_cb') and self.fin_imca_cb.isChecked():
                        domain_actions.append("IMCA referral")
                    # Radio button group - only one can be selected
                    if hasattr(self, 'fin_guardianship_rb') and self.fin_guardianship_rb.isChecked():
                        domain_actions.append("Guardianship")
                    elif hasattr(self, 'fin_appointeeship_rb') and self.fin_appointeeship_rb.isChecked():
                        domain_actions.append("Appointeeship")
                    elif hasattr(self, 'fin_informal_rb') and self.fin_informal_rb.isChecked():
                        domain_actions.append("Informal Appointeeship")

                elif domain == "residence":
                    if hasattr(self, 'res_best_interest_cb') and self.res_best_interest_cb.isChecked():
                        domain_actions.append("Best Interest Meeting")
                    if hasattr(self, 'res_imca_cb') and self.res_imca_cb.isChecked():
                        domain_actions.append("IMCA referral")
                    if hasattr(self, 'res_dols_cb') and self.res_dols_cb.isChecked():
                        domain_actions.append("DOLS")
                    if hasattr(self, 'res_cop_cb') and self.res_cop_cb.isChecked():
                        domain_actions.append("Court of Protection")

                elif domain == "self-care":
                    if hasattr(self, 'self_best_interest_cb') and self.self_best_interest_cb.isChecked():
                        domain_actions.append("Best Interest Meeting")
                    if hasattr(self, 'self_imca_cb') and self.self_imca_cb.isChecked():
                        domain_actions.append("IMCA referral")
                    if hasattr(self, 'self_dols_cb') and self.self_dols_cb.isChecked():
                        domain_actions.append("DOLS")
                    if hasattr(self, 'self_cop_cb') and self.self_cop_cb.isChecked():
                        domain_actions.append("Court of Protection")

                if domain_actions:
                    if len(domain_actions) == 1:
                        parts.append(f"Action: {domain_actions[0]}.")
                    else:
                        items = ", ".join(domain_actions[:-1]) + f" and {domain_actions[-1]}"
                        parts.append(f"Actions: {items}.")
        else:
            parts.append("Capacity assessment has not been formally documented.")

        new_generated = " ".join(parts)
        self._update_text_preserving_additions(self.capacity_text, new_generated, "capacity")

    # ================================================================
    # SECTION 8: Progress
    # ================================================================
    def _build_section_8_progress(self):
        frame = self._create_section_frame("Progress", "8", "#7c3aed")
        layout = frame.layout()

        split = QHBoxLayout()
        split.setSpacing(16)

        self.progress_text = self._create_text_edit("", 200)
        split.addWidget(self.progress_text, 3)

        # Right panel with progress area sliders
        right_panel = QFrame()
        right_panel.setStyleSheet("QFrame { background: #f5f3ff; border-radius: 8px; }")
        right_panel.setFixedWidth(420)
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(12, 10, 12, 10)
        right_layout.setSpacing(4)

        # Progress areas header
        prog_header = QLabel("Progress Areas:")
        prog_header.setStyleSheet("font-size: 16px; font-weight: 600; color: #5b21b6;")
        right_layout.addWidget(prog_header)

        # Progress sliders with specific options
        self.progress_sliders = {}
        progress_options = [
            ("mental_state", "Mental State", [
                "Unsettled", "Often unsettled", "Unsettled at times",
                "Stable", "Some improvement", "Significant improvement",
                "Symptom free with no concerns"
            ]),
            ("insight", "Insight", [
                "Remains limited", "Mostly absent but some insight",
                "Mild insight", "Moderate insight", "Good insight", "Full insight"
            ]),
            ("engagement", "Engagement with Treatment", [
                "Nil", "Some", "Partial", "Good", "Very good", "Full"
            ]),
            ("risk_reduction", "Risk Reduction Work", [
                "Nil", "Started", "In process", "Good engagement",
                "Almost completed", "Completed"
            ]),
            ("leave", "Leave", [
                "No leave", "Escorted", "Unescorted", "Overnight"
            ]),
            ("discharge_planning", "Discharge Planning", [
                "Not started", "Early stages", "In progress",
                "Almost completed", "Completed"
            ]),
        ]

        for key, label, options in progress_options:
            # Slider frame
            slider_frame = QFrame()
            slider_frame.setStyleSheet("QFrame { background: #ede9fe; border-radius: 4px; }")
            slider_layout = QVBoxLayout(slider_frame)
            slider_layout.setContentsMargins(8, 4, 8, 4)
            slider_layout.setSpacing(2)

            # Label
            lbl = QLabel(label)
            lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #5b21b6;")
            slider_layout.addWidget(lbl)

            # Slider row
            slider_row = QHBoxLayout()
            slider_row.setSpacing(6)

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setMinimum(0)
            slider.setMaximum(len(options) - 1)
            slider.setValue(0)
            slider.setFixedWidth(160)
            slider.valueChanged.connect(lambda v, k=key: self._on_progress_slider_changed(k))
            slider_row.addWidget(slider)

            val_lbl = QLabel(options[0])
            val_lbl.setStyleSheet("font-size: 14px; color: #374151; min-width: 160px;")
            slider_row.addWidget(val_lbl)
            slider_row.addStretch()

            slider_layout.addLayout(slider_row)

            self.progress_sliders[key] = {
                "slider": slider,
                "value_label": val_lbl,
                "options": options,
                "label": label
            }

            # Add secondary leave usage and concerns sliders (hidden by default)
            if key == "leave":
                self.leave_usage_widget = QWidget()
                leave_usage_layout = QVBoxLayout(self.leave_usage_widget)
                leave_usage_layout.setContentsMargins(0, 2, 0, 0)
                leave_usage_layout.setSpacing(4)

                # Leave Use slider
                usage_lbl = QLabel("Leave Use")
                usage_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #7c3aed;")
                leave_usage_layout.addWidget(usage_lbl)

                usage_row = QHBoxLayout()
                usage_row.setSpacing(6)

                self.leave_usage_slider = NoWheelSlider(Qt.Orientation.Horizontal)
                self.leave_usage_slider.setMinimum(0)
                self.leave_usage_slider.setMaximum(4)
                self.leave_usage_slider.setValue(2)
                self.leave_usage_slider.setFixedWidth(160)
                self.leave_usage_slider.valueChanged.connect(self._on_leave_usage_changed)
                usage_row.addWidget(self.leave_usage_slider)

                self.leave_usage_options = ["Intermittent", "Variable", "Regular", "Frequent", "Excellent"]
                self.leave_usage_label = QLabel(self.leave_usage_options[2])
                self.leave_usage_label.setStyleSheet("font-size: 14px; color: #374151; min-width: 160px;")
                usage_row.addWidget(self.leave_usage_label)
                usage_row.addStretch()
                leave_usage_layout.addLayout(usage_row)

                # Leave Concerns slider
                concerns_lbl = QLabel("Leave Concerns")
                concerns_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #dc2626;")
                leave_usage_layout.addWidget(concerns_lbl)

                concerns_row = QHBoxLayout()
                concerns_row.setSpacing(6)

                self.leave_concerns_slider = NoWheelSlider(Qt.Orientation.Horizontal)
                self.leave_concerns_slider.setMinimum(0)
                self.leave_concerns_slider.setMaximum(4)
                self.leave_concerns_slider.setValue(0)
                self.leave_concerns_slider.setFixedWidth(160)
                self.leave_concerns_slider.valueChanged.connect(self._on_leave_concerns_changed)
                concerns_row.addWidget(self.leave_concerns_slider)

                self.leave_concerns_options = ["No", "Mild", "Some", "Several", "Significant"]
                self.leave_concerns_label = QLabel(self.leave_concerns_options[0])
                self.leave_concerns_label.setStyleSheet("font-size: 14px; color: #374151; min-width: 160px;")
                concerns_row.addWidget(self.leave_concerns_label)
                concerns_row.addStretch()
                leave_usage_layout.addLayout(concerns_row)

                self.leave_usage_widget.setVisible(False)
                slider_layout.addWidget(self.leave_usage_widget)

            right_layout.addWidget(slider_frame)

        # Discharge application section
        right_layout.addSpacing(6)
        discharge_lbl = QLabel("Discharge applications:")
        discharge_lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #5b21b6;")
        right_layout.addWidget(discharge_lbl)

        self.discharge_app_group = QButtonGroup(self)
        self.discharge_app_yes = QRadioButton("Yes - applications made")
        self.discharge_app_no = QRadioButton("No applications")
        for rb in [self.discharge_app_yes, self.discharge_app_no]:
            rb.setStyleSheet("font-size: 14px; color: #374151;")
            self.discharge_app_group.addButton(rb)
            rb.toggled.connect(self._update_progress_text)
            right_layout.addWidget(rb)

        right_layout.addStretch()
        split.addWidget(right_panel)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def _on_progress_slider_changed(self, key):
        """Update slider label and text when slider changes."""
        data = self.progress_sliders[key]
        value = data["slider"].value()
        data["value_label"].setText(data["options"][value])

        # Show/hide leave usage slider based on leave type
        if key == "leave":
            self.leave_usage_widget.setVisible(value > 0)

        self._update_progress_text()

    def _on_leave_usage_changed(self, value):
        """Update leave usage label."""
        self.leave_usage_label.setText(self.leave_usage_options[value])
        self._update_progress_text()

    def _on_leave_concerns_changed(self, value):
        """Update leave concerns label."""
        self.leave_concerns_label.setText(self.leave_concerns_options[value])
        self._update_progress_text()

    def _update_progress_text(self):
        """Update progress text based on slider selections."""
        parts = []

        # Get pronouns from patient gender if available
        gender = getattr(self, '_patient_gender', 'male')
        if gender == 'female':
            subj, obj, pos = "She", "her", "her"
        elif gender == 'male':
            subj, obj, pos = "He", "him", "his"
        else:
            subj, obj, pos = "They", "them", "their"

        # Mental state
        ms_idx = self.progress_sliders["mental_state"]["slider"].value()
        ms_phrases = [
            "mental state has been unsettled",
            "mental state has often been unsettled",
            "mental state has been unsettled at times",
            "mental state has remained stable",
            "mental state has shown some improvement",
            "mental state has shown significant improvement",
            "mental state has been symptom free with no concerns"
        ]
        parts.append(f"Over the last 12 months, {pos} {ms_phrases[ms_idx]}.")

        # Insight
        ins_idx = self.progress_sliders["insight"]["slider"].value()
        ins_phrases = [
            f"{subj} has displayed limited insight into {pos} needs and {pos} illness",
            f"{subj} has displayed mostly absent but some insight into {pos} needs and {pos} illness",
            f"{subj} has displayed mild insight into {pos} needs and {pos} illness",
            f"{subj} has displayed moderate insight into {pos} needs and {pos} illness",
            f"{subj} has displayed good insight into {pos} needs and {pos} illness",
            f"{subj} has displayed full insight into {pos} needs and {pos} illness"
        ]
        parts.append(f"{ins_phrases[ins_idx]}.")

        # Engagement and Leave combined
        eng_idx = self.progress_sliders["engagement"]["slider"].value()
        eng_phrases = ["nil", "some", "partial", "good", "very good", "full"]

        lv_idx = self.progress_sliders["leave"]["slider"].value()
        if lv_idx == 0:
            parts.append(f"{pos.capitalize()} engagement with treatment has been {eng_phrases[eng_idx]} overall and {subj.lower()} has not been taking any leave.")
        else:
            lv_types = ["", "escorted", "unescorted", "overnight"]
            usage_idx = self.leave_usage_slider.value()
            usage_phrases = ["intermittent", "variable", "regular", "frequent", "excellent"]
            concerns_idx = self.leave_concerns_slider.value()
            concerns_phrases = ["no", "mild", "some", "several", "significant"]
            if concerns_idx == 0:
                parts.append(f"{pos.capitalize()} engagement with treatment has been {eng_phrases[eng_idx]} overall and {subj.lower()} has been taking {lv_types[lv_idx]} leave on a {usage_phrases[usage_idx]} basis with no concerns.")
            else:
                parts.append(f"{pos.capitalize()} engagement with treatment has been {eng_phrases[eng_idx]} overall and {subj.lower()} has been taking {lv_types[lv_idx]} leave on a {usage_phrases[usage_idx]} basis with {concerns_phrases[concerns_idx]} concerns.")

        # Risk reduction work
        rr_idx = self.progress_sliders["risk_reduction"]["slider"].value()
        rr_phrases = [
            f"{subj} has not yet engaged in risk reduction work",
            f"{subj} has started risk reduction work",
            f"Risk reduction work is in process",
            f"{subj} has shown good engagement with risk reduction work",
            f"Risk reduction work is almost completed",
            f"Risk reduction work has been completed"
        ]
        parts.append(f"{rr_phrases[rr_idx]}.")

        # Discharge planning
        dp_idx = self.progress_sliders["discharge_planning"]["slider"].value()
        dp_phrases = [
            "discharge planning has not yet started",
            "discharge planning is at an early stage",
            "discharge planning is in progress",
            "discharge planning is almost completed",
            "discharge planning has been completed"
        ]
        parts.append(f"Currently {dp_phrases[dp_idx]}.")

        # Discharge applications
        if self.discharge_app_yes.isChecked():
            parts.append("During this period applications have been made for discharge.")
        elif self.discharge_app_no.isChecked():
            parts.append("During this period no applications have been made for discharge.")

        new_generated = " ".join(parts)
        self._update_text_preserving_additions(self.progress_text, new_generated, "progress")

    # ================================================================
    # SECTION 9: Managing Risk - Risk Factors (matches Section 5)
    # ================================================================
    RISK_TYPES = [
        ("violence_others", "Violence to others"),
        ("violence_property", "Violence to property"),
        ("verbal_aggression", "Verbal aggression"),
        ("substance_misuse", "Substance misuse"),
        ("self_harm", "Self harm"),
        ("self_neglect", "Self neglect"),
        ("stalking", "Stalking"),
        ("threatening_behaviour", "Threatening behaviour"),
        ("sexually_inappropriate", "Sexually inappropriate behaviour"),
        ("vulnerability", "Vulnerability"),
        ("bullying_victimisation", "Bullying/victimisation"),
        ("absconding", "Absconding/AWOL"),
        ("reoffending", "Reoffending"),
    ]

    def _build_section_9_managing_risk(self):
        frame = self._create_section_frame("Managing Risk", "9", "#dc2626")
        layout = frame.layout()

        info_lbl = QLabel("It is important for the Secretary of State to understand the clinical assessment of risk.")
        info_lbl.setWordWrap(True)
        info_lbl.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        layout.addWidget(info_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.risk_factors_text = self._create_text_edit("", 280)
        split.addWidget(self.risk_factors_text, 3)

        # Right panel with Current and Historical risk sections
        right_panel = QFrame()
        right_panel.setStyleSheet("QFrame { background: #fef2f2; border-radius: 8px; }")
        right_panel.setFixedWidth(420)
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(10, 10, 10, 10)
        right_layout.setSpacing(8)

        # Scrollable content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(8)

        # Current Risk section
        current_header = QLabel("Current Risk")
        current_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #991b1b; background: #fee2e2; padding: 4px 8px; border-radius: 4px;")
        scroll_layout.addWidget(current_header)

        self._current_risk_widgets = {}
        for key, label in self.RISK_TYPES:
            row_widget = self._create_risk_row(key, label, "current")
            scroll_layout.addWidget(row_widget)

        scroll_layout.addSpacing(8)

        # Historical Risk section
        hist_header = QLabel("Historical Risk")
        hist_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #806000; background: #fef3c7; padding: 4px 8px; border-radius: 4px;")
        scroll_layout.addWidget(hist_header)

        self._historical_risk_widgets = {}
        for key, label in self.RISK_TYPES:
            row_widget = self._create_risk_row(key, label, "historical")
            scroll_layout.addWidget(row_widget)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_content)
        right_layout.addWidget(scroll)

        split.addWidget(right_panel)
        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def _create_risk_row(self, key, label, risk_type):
        """Create a risk row with checkbox and severity slider."""
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container_layout = QVBoxLayout(container)
        container_layout.setContentsMargins(0, 0, 0, 0)
        container_layout.setSpacing(2)

        cb = QCheckBox(label)
        cb.setStyleSheet("font-size: 15px; color: #374151;")
        cb.toggled.connect(self._update_risk_text)
        container_layout.addWidget(cb)

        # Severity slider (hidden by default)
        slider_widget = QWidget()
        slider_widget.setStyleSheet("background: transparent;")
        slider_layout = QHBoxLayout(slider_widget)
        slider_layout.setContentsMargins(20, 0, 0, 4)
        slider_layout.setSpacing(6)

        slider_lbl = QLabel("Severity:")
        slider_lbl.setStyleSheet("font-size: 14px; color: #6b7280;")
        slider_layout.addWidget(slider_lbl)

        slider = NoWheelSlider(Qt.Orientation.Horizontal)
        slider.setMinimum(1)
        slider.setMaximum(3)
        slider.setValue(2)
        slider.setFixedWidth(80)
        slider.valueChanged.connect(self._update_risk_text)
        slider_layout.addWidget(slider)

        level_lbl = QLabel("Medium")
        level_lbl.setStyleSheet("font-size: 14px; color: #374151; font-weight: 500; min-width: 50px;")
        slider.valueChanged.connect(lambda v, l=level_lbl: l.setText(["Low", "Medium", "High"][v-1]))
        slider_layout.addWidget(level_lbl)
        slider_layout.addStretch()

        slider_widget.setVisible(False)
        container_layout.addWidget(slider_widget)

        cb.toggled.connect(lambda checked, sw=slider_widget: sw.setVisible(checked))

        # Store widgets
        widgets_dict = self._current_risk_widgets if risk_type == "current" else self._historical_risk_widgets
        widgets_dict[key] = {
            "checkbox": cb,
            "slider": slider,
            "slider_widget": slider_widget,
            "level_label": level_lbl
        }

        return container

    def _update_risk_text(self):
        """Update risk factors text based on selections."""
        parts = []

        # Current risks
        current_risks = []
        for key, label in self.RISK_TYPES:
            widgets = self._current_risk_widgets.get(key, {})
            if widgets.get("checkbox") and widgets["checkbox"].isChecked():
                level = ["low", "medium", "high"][widgets["slider"].value() - 1]
                current_risks.append(f"{label.lower()} ({level})")

        if current_risks:
            parts.append("Current risks: " + ", ".join(current_risks) + ".")
        else:
            parts.append("No current risks identified.")

        # Historical risks
        historical_risks = []
        for key, label in self.RISK_TYPES:
            widgets = self._historical_risk_widgets.get(key, {})
            if widgets.get("checkbox") and widgets["checkbox"].isChecked():
                level = ["low", "medium", "high"][widgets["slider"].value() - 1]
                historical_risks.append(f"{label.lower()} ({level})")

        if historical_risks:
            parts.append("Historical risks: " + ", ".join(historical_risks) + ".")
        else:
            parts.append("No historical risks identified.")

        new_generated = " ".join(parts)
        self._update_text_preserving_additions(self.risk_factors_text, new_generated, "risk_factors")

    # ================================================================
    # SECTION 10: How Risks Addressed
    # ================================================================
    def _build_section_10_risk_addressed(self):
        frame = self._create_section_frame("Risks Addressed", "10", "#dc2626")
        layout = frame.layout()

        prompt_lbl = QLabel("Describe how these risks have been addressed in the last 12 months:")
        prompt_lbl.setWordWrap(True)
        prompt_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(prompt_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.risk_addressed_text = self._create_text_edit("", 180)
        split.addWidget(self.risk_addressed_text, 3)

        prompts_frame, self.risk_addressed_checkboxes = self._create_prompts_frame([
            "Describe the progress the patient (either in hospital and/or the community) has made and any issues of concern",
            "What is the team's current understanding of the factors underpinning the index offence and previous dangerous behaviour",
            "What are the patient's current attitudes to the index offence, other dangerous behaviour and any previous victims",
            "Please advise if the patient has been referred to Prevent and if so, the outcome of that referral"
        ], "#dc2626")
        prompts_frame.setFixedWidth(400)
        split.addWidget(prompts_frame)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    # ================================================================
    # SECTION 11: Abscond / Escape
    # ================================================================
    def _build_section_11_abscond(self):
        frame = self._create_section_frame("Abscond / Escape", "11", "#dc2626")
        layout = frame.layout()

        split = QHBoxLayout()
        split.setSpacing(16)

        self.abscond_text = self._create_text_edit("", 120)
        split.addWidget(self.abscond_text, 3)

        # Right panel with AWOL-style controls
        right_panel = QFrame()
        right_panel.setStyleSheet("QFrame { background: #fef2f2; border-radius: 8px; }")
        right_panel.setFixedWidth(340)
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(12, 10, 12, 10)
        right_layout.setSpacing(6)

        # AWOL status header
        awol_header = QLabel("AWOL / Abscond History:")
        awol_header.setStyleSheet("font-size: 15px; font-weight: 600; color: #991b1b;")
        right_layout.addWidget(awol_header)

        # Yes/No toggle buttons
        btn_row = QHBoxLayout()
        btn_row.setSpacing(8)

        self.awol_yes_btn = QPushButton("Yes - AWOL concerns")
        self.awol_no_btn = QPushButton("No AWOL concerns")

        for btn in [self.awol_yes_btn, self.awol_no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton { background: #fee2e2; color: #991b1b; border: 1px solid #fecaca;
                              border-radius: 6px; padding: 6px 10px; font-size: 14px; font-weight: 500; }
                QPushButton:checked { background: #14b8a6; color: white; border-color: #14b8a6; }
                QPushButton:hover { background: #fecaca; }
                QPushButton:checked:hover { background: #0d9488; }
            """)
            btn_row.addWidget(btn)

        self.awol_yes_btn.clicked.connect(lambda: self._on_awol_toggled(True))
        self.awol_no_btn.clicked.connect(lambda: self._on_awol_toggled(False))

        right_layout.addLayout(btn_row)

        # Details section (shown when Yes selected)
        self.awol_details_frame = QFrame()
        details_layout = QVBoxLayout(self.awol_details_frame)
        details_layout.setContentsMargins(0, 8, 0, 0)
        details_layout.setSpacing(4)

        details_lbl = QLabel("Incident details:")
        details_lbl.setStyleSheet("font-size: 14px; font-weight: 600; color: #991b1b;")
        details_layout.addWidget(details_lbl)

        # Incident type checkboxes
        self.awol_types = {}
        awol_options = [
            ("escaped", "Escaped from escort"),
            ("absconded", "Absconded from ward"),
            ("failed_return", "Failed to return from leave"),
            ("went_missing", "Went missing"),
        ]

        for key, label in awol_options:
            cb = QCheckBox(label)
            cb.setStyleSheet("font-size: 14px; color: #374151;")
            cb.stateChanged.connect(self._update_abscond_text)
            details_layout.addWidget(cb)
            self.awol_types[key] = cb

        # Additional details input
        self.awol_details_input = QLineEdit()
        self.awol_details_input.setPlaceholderText("Additional details (dates, circumstances)...")
        self.awol_details_input.setStyleSheet("font-size: 14px; padding: 4px;")
        self.awol_details_input.textChanged.connect(self._update_abscond_text)
        details_layout.addWidget(self.awol_details_input)

        self.awol_details_frame.setVisible(False)
        right_layout.addWidget(self.awol_details_frame)

        right_layout.addStretch()
        split.addWidget(right_panel)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def _on_awol_toggled(self, is_yes):
        """Handle AWOL Yes/No toggle."""
        if is_yes:
            self.awol_yes_btn.setChecked(True)
            self.awol_no_btn.setChecked(False)
            self.awol_details_frame.setVisible(True)
        else:
            self.awol_yes_btn.setChecked(False)
            self.awol_no_btn.setChecked(True)
            self.awol_details_frame.setVisible(False)
            # Clear selections
            for cb in self.awol_types.values():
                cb.setChecked(False)
            self.awol_details_input.clear()
        self._update_abscond_text()

    def _update_abscond_text(self):
        """Update abscond text based on selections."""
        new_generated = ""

        if self.awol_no_btn.isChecked():
            new_generated = (
                "There have been no occasions on which the patient has been absent without leave, "
                "escaped, or failed to return from authorised leave."
            )
        elif self.awol_yes_btn.isChecked():
            parts = []
            incident_types = []

            type_descriptions = {
                "escaped": "escaped from escort",
                "absconded": "absconded from the ward",
                "failed_return": "failed to return from authorised leave",
                "went_missing": "went missing",
            }

            for key, cb in self.awol_types.items():
                if cb.isChecked():
                    incident_types.append(type_descriptions.get(key, cb.text()))

            if incident_types:
                if len(incident_types) == 1:
                    parts.append(f"The patient has {incident_types[0]}.")
                else:
                    items = ", ".join(incident_types[:-1]) + f" and {incident_types[-1]}"
                    parts.append(f"The patient has {items}.")

            details = self.awol_details_input.text().strip()
            if details:
                parts.append(details)

            if parts:
                new_generated = " ".join(parts)
            else:
                new_generated = "The patient has had AWOL concerns (details to be provided)."

        if new_generated:
            self._update_text_preserving_additions(self.abscond_text, new_generated, "abscond")

    # ================================================================
    # SECTION 12: MAPPA
    # ================================================================
    def _build_section_12_mappa(self):
        frame = self._create_section_frame("MAPPA", "12", "#ea580c")
        layout = frame.layout()

        info_lbl = QLabel("It is the Responsible Clinician's responsibility to refer patients to Multi-Agency Public Protection Arrangements (MAPPA) where applicable:")
        info_lbl.setWordWrap(True)
        info_lbl.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        layout.addWidget(info_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.mappa_text = self._create_text_edit("", 150)
        split.addWidget(self.mappa_text, 3)

        # Right panel with MAPPA category/level radio buttons
        right_panel = QFrame()
        right_panel.setFixedWidth(350)
        right_panel.setStyleSheet("QFrame { background: #fff7ed; border-radius: 8px; }")
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(12, 12, 12, 12)
        right_layout.setSpacing(10)

        # MAPPA Category
        cat_lbl = QLabel("MAPPA Category")
        cat_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #ea580c;")
        right_layout.addWidget(cat_lbl)

        self.mappa_cat_group = QButtonGroup(self)
        self.mappa_cat_1 = QRadioButton("Category 1 - Registered Sex Offender")
        self.mappa_cat_2 = QRadioButton("Category 2 - Violent Offender")
        self.mappa_cat_3 = QRadioButton("Category 3 - Other Dangerous Offender")
        self.mappa_cat_na = QRadioButton("Not applicable / Not referred")
        self.mappa_cat_na.setChecked(True)

        for i, rb in enumerate([self.mappa_cat_na, self.mappa_cat_1, self.mappa_cat_2, self.mappa_cat_3]):
            rb.setStyleSheet("font-size: 16px; color: #374151;")
            rb.toggled.connect(self._update_mappa_preview)
            self.mappa_cat_group.addButton(rb, i)
            right_layout.addWidget(rb)

        right_layout.addSpacing(8)

        # MAPPA Level
        level_lbl = QLabel("MAPPA Level")
        level_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #ea580c;")
        right_layout.addWidget(level_lbl)

        self.mappa_level_group = QButtonGroup(self)
        self.mappa_level_1 = QRadioButton("Level 1 - Ordinary Agency Management")
        self.mappa_level_2 = QRadioButton("Level 2 - Active Multi-Agency Management")
        self.mappa_level_3 = QRadioButton("Level 3 - Senior Management Oversight")
        self.mappa_level_na = QRadioButton("N/A")
        self.mappa_level_na.setChecked(True)

        for i, rb in enumerate([self.mappa_level_na, self.mappa_level_1, self.mappa_level_2, self.mappa_level_3]):
            rb.setStyleSheet("font-size: 16px; color: #374151;")
            rb.toggled.connect(self._update_mappa_preview)
            self.mappa_level_group.addButton(rb, i)
            right_layout.addWidget(rb)

        right_layout.addSpacing(8)

        # Referral date
        ref_lbl = QLabel("Date of Referral")
        ref_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #ea580c;")
        right_layout.addWidget(ref_lbl)

        self.mappa_referral_date = QLineEdit()
        self.mappa_referral_date.setPlaceholderText("DD/MM/YYYY")
        self.mappa_referral_date.setStyleSheet("font-size: 16px; padding: 6px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.mappa_referral_date.textChanged.connect(self._update_mappa_preview)
        right_layout.addWidget(self.mappa_referral_date)

        right_layout.addStretch()
        split.addWidget(right_panel)

        layout.addLayout(split)

        # MAPPA Coordinator details
        coord_lbl = QLabel("Please give the name and contact details of the MAPPA coordinator:")
        coord_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(coord_lbl)
        self.mappa_coordinator = self._create_line_edit()
        layout.addWidget(self.mappa_coordinator)

        self.form_layout.addWidget(frame)

    def _update_mappa_preview(self):
        """Update MAPPA text based on radio selections."""
        parts = []

        # Category
        if self.mappa_cat_1.isChecked():
            parts.append("MAPPA Category 1 (Registered Sex Offender).")
        elif self.mappa_cat_2.isChecked():
            parts.append("MAPPA Category 2 (Violent Offender).")
        elif self.mappa_cat_3.isChecked():
            parts.append("MAPPA Category 3 (Other Dangerous Offender).")
        elif self.mappa_cat_na.isChecked():
            parts.append("Patient is not currently referred to MAPPA.")

        # Level (only if category is selected)
        if not self.mappa_cat_na.isChecked():
            if self.mappa_level_1.isChecked():
                parts.append("Managed at Level 1 (Ordinary Agency Management).")
            elif self.mappa_level_2.isChecked():
                parts.append("Managed at Level 2 (Active Multi-Agency Management).")
            elif self.mappa_level_3.isChecked():
                parts.append("Managed at Level 3 (Senior Management Oversight).")

            # Referral date
            ref_date = self.mappa_referral_date.text().strip()
            if ref_date:
                parts.append(f"Referred to MAPPA on {ref_date}.")

        self.mappa_text.setPlainText(" ".join(parts))

    # ================================================================
    # SECTION 13: Victims
    # ================================================================
    def _build_section_13_victims(self):
        frame = self._create_section_frame("Victims", "13", "#be185d")
        layout = frame.layout()

        info_lbl = QLabel("Not all victims will be registered with the Victim Liaison Scheme. It is MHCS policy to take into account any information provided by victims to ensure they feel adequately protected. Multi-Disciplinary teams should try to ensure victims are not inadvertently put at risk by the actions of their patients.")
        info_lbl.setWordWrap(True)
        info_lbl.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        layout.addWidget(info_lbl)

        vlo_lbl = QLabel("Victim Liaison Officer (VLO) name and contact details:")
        vlo_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(vlo_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.victims_text = self._create_text_edit("", 140)
        split.addWidget(self.victims_text, 3)

        prompts_frame, self.victims_checkboxes = self._create_prompts_frame([
            "Please provide full contact details",
            "Give date of last discussion/contact with VLO",
            "Have there been any victim-related concerns or contact with the VLO in the last 12 months"
        ], "#be185d")
        prompts_frame.setFixedWidth(380)
        split.addWidget(prompts_frame)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    # ================================================================
    # SECTION 14: Community Leave - Leave Types
    # ================================================================
    def _build_section_14_community_leave(self):
        frame = self._create_section_frame("Community Leave", "14", "#0891b2")
        layout = frame.layout()

        use_lbl = QLabel("Use of Community Leave")
        use_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #0891b2;")
        layout.addWidget(use_lbl)

        prompt_lbl = QLabel("Community leave taken (please mark X in box):")
        prompt_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        layout.addWidget(prompt_lbl)

        grid = QGridLayout()
        grid.setSpacing(12)

        self.leave_compassionate_day = QCheckBox("Compassionate (day)")
        self.leave_compassionate_overnight = QCheckBox("Compassionate (overnight)")
        self.leave_medical_day = QCheckBox("Medical (day)")
        self.leave_medical_overnight = QCheckBox("Medical (overnight)")
        self.leave_escorted_day = QCheckBox("Escorted (day)")
        self.leave_escorted_overnight = QCheckBox("Escorted (overnight)")
        self.leave_unescorted_day = QCheckBox("Unescorted community (day)")
        self.leave_unescorted_overnight = QCheckBox("Unescorted community (overnight)")
        self.leave_long_term_escorted = QCheckBox("Long Term Escorted Leave of Absence")

        for cb in [self.leave_compassionate_day, self.leave_compassionate_overnight,
                   self.leave_medical_day, self.leave_medical_overnight,
                   self.leave_escorted_day, self.leave_escorted_overnight,
                   self.leave_unescorted_day, self.leave_unescorted_overnight,
                   self.leave_long_term_escorted]:
            cb.setStyleSheet("font-size: 17px; color: #374151;")

        grid.addWidget(self.leave_compassionate_day, 0, 0)
        grid.addWidget(self.leave_compassionate_overnight, 0, 1)
        grid.addWidget(self.leave_medical_day, 1, 0)
        grid.addWidget(self.leave_medical_overnight, 1, 1)
        grid.addWidget(self.leave_escorted_day, 2, 0)
        grid.addWidget(self.leave_escorted_overnight, 2, 1)
        grid.addWidget(self.leave_unescorted_day, 3, 0)
        grid.addWidget(self.leave_unescorted_overnight, 3, 1)
        grid.addWidget(self.leave_long_term_escorted, 4, 0, 1, 2)

        layout.addLayout(grid)
        self.form_layout.addWidget(frame)

    # ================================================================
    # SECTION 14b: Leave Details (part of Leave Report)
    # ================================================================
    def _build_section_14b_leave_details(self):
        frame = self._create_section_frame("Leave Details", "14", "#0891b2")
        layout = frame.layout()

        split = QHBoxLayout()
        split.setSpacing(16)

        self.leave_details_text = self._create_text_edit("", 220)
        split.addWidget(self.leave_details_text, 3)

        # Right panel with leave controls
        right_panel = QFrame()
        right_panel.setFixedWidth(400)
        right_panel.setStyleSheet("QFrame { background: #ecfeff; border-radius: 8px; }")
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(10, 10, 10, 10)
        right_layout.setSpacing(8)

        # 1. Leave taken Yes/No
        leave_taken_row = QHBoxLayout()
        leave_taken_row.setSpacing(8)
        leave_taken_lbl = QLabel("Leave taken in last 12 months:")
        leave_taken_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #0891b2;")
        leave_taken_row.addWidget(leave_taken_lbl)

        self.leave_det_taken_group = QButtonGroup(self)
        self.leave_det_taken_yes = QRadioButton("Yes")
        self.leave_det_taken_no = QRadioButton("No")
        self.leave_det_taken_no.setChecked(True)
        for rb in [self.leave_det_taken_yes, self.leave_det_taken_no]:
            rb.setStyleSheet("font-size: 15px; color: #374151;")
            self.leave_det_taken_group.addButton(rb)
            leave_taken_row.addWidget(rb)
        leave_taken_row.addStretch()
        self.leave_det_taken_yes.toggled.connect(self._on_leave_details_taken_toggled)
        self.leave_det_taken_no.toggled.connect(self._update_leave_details_text)
        right_layout.addLayout(leave_taken_row)

        # Leave types container (shown when Yes)
        self.leave_det_types_frame = QFrame()
        leave_types_layout = QVBoxLayout(self.leave_det_types_frame)
        leave_types_layout.setContentsMargins(0, 4, 0, 0)
        leave_types_layout.setSpacing(4)

        types_lbl = QLabel("Leave types:")
        types_lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #0891b2;")
        leave_types_layout.addWidget(types_lbl)

        # Leave type checkboxes with sliders
        self.leave_det_type_widgets = {}
        leave_types = [
            ("grounds", "Grounds"),
            ("local", "Local community"),
            ("community", "Community"),
            ("overnight", "Overnight")
        ]

        for type_key, type_label in leave_types:
            type_frame = QFrame()
            type_frame.setStyleSheet("QFrame { background: #cffafe; border-radius: 4px; }")
            type_layout = QVBoxLayout(type_frame)
            type_layout.setContentsMargins(6, 4, 6, 4)
            type_layout.setSpacing(3)

            cb = QCheckBox(type_label)
            cb.setStyleSheet("font-size: 15px; color: #374151; font-weight: 500;")
            cb.toggled.connect(lambda checked, k=type_key: self._on_leave_det_type_toggled(k, checked))
            type_layout.addWidget(cb)

            # Sliders container (hidden until checkbox checked)
            sliders_widget = QWidget()
            sliders_layout = QVBoxLayout(sliders_widget)
            sliders_layout.setContentsMargins(10, 0, 0, 0)
            sliders_layout.setSpacing(2)

            # Frequency slider
            freq_row = QHBoxLayout()
            freq_row.setSpacing(4)
            freq_lbl = QLabel("Frequency:")
            freq_lbl.setStyleSheet("font-size: 14px; color: #6b7280;")
            freq_row.addWidget(freq_lbl)

            freq_slider = NoWheelSlider(Qt.Orientation.Horizontal)
            freq_slider.setMinimum(1)
            freq_slider.setMaximum(7)
            freq_slider.setValue(1)
            freq_slider.setFixedWidth(100)
            freq_slider.valueChanged.connect(lambda v, k=type_key: self._on_leave_det_slider_changed(k))
            freq_row.addWidget(freq_slider)

            freq_val = QLabel("1x/week")
            freq_val.setStyleSheet("font-size: 14px; color: #374151; min-width: 55px;")
            freq_row.addWidget(freq_val)
            freq_row.addStretch()
            sliders_layout.addLayout(freq_row)

            # Duration slider
            dur_row = QHBoxLayout()
            dur_row.setSpacing(4)
            dur_lbl = QLabel("Duration:")
            dur_lbl.setStyleSheet("font-size: 14px; color: #6b7280;")
            dur_row.addWidget(dur_lbl)

            dur_slider = NoWheelSlider(Qt.Orientation.Horizontal)
            dur_slider.setMinimum(1)
            dur_slider.setMaximum(16)
            dur_slider.setValue(1)
            dur_slider.setFixedWidth(100)
            dur_slider.valueChanged.connect(lambda v, k=type_key: self._on_leave_det_slider_changed(k))
            dur_row.addWidget(dur_slider)

            dur_val = QLabel("30 mins")
            dur_val.setStyleSheet("font-size: 14px; color: #374151; min-width: 55px;")
            dur_row.addWidget(dur_val)
            dur_row.addStretch()
            sliders_layout.addLayout(dur_row)

            sliders_widget.setVisible(False)
            type_layout.addWidget(sliders_widget)

            leave_types_layout.addWidget(type_frame)

            self.leave_det_type_widgets[type_key] = {
                "checkbox": cb,
                "sliders_widget": sliders_widget,
                "freq_slider": freq_slider,
                "freq_val": freq_val,
                "dur_slider": dur_slider,
                "dur_val": dur_val
            }

        self.leave_det_types_frame.setVisible(False)
        right_layout.addWidget(self.leave_det_types_frame)

        # 2. Leave to attend
        attend_lbl = QLabel("Leave to attend:")
        attend_lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #0891b2; margin-top: 4px;")
        right_layout.addWidget(attend_lbl)

        self.leave_det_attend = {}
        for att_key, att_label in [("court", "Court"), ("medical", "Medical appointments"), ("dental", "Dental appointments")]:
            cb = QCheckBox(att_label)
            cb.setStyleSheet("font-size: 15px; color: #374151;")
            cb.toggled.connect(self._update_leave_details_text)
            right_layout.addWidget(cb)
            self.leave_det_attend[att_key] = cb

        # 3. Leave suspensions Yes/No
        susp_row = QHBoxLayout()
        susp_row.setSpacing(8)
        susp_lbl = QLabel("Leave suspended:")
        susp_lbl.setStyleSheet("font-size: 15px; font-weight: 600; color: #0891b2; margin-top: 4px;")
        susp_row.addWidget(susp_lbl)

        self.leave_det_susp_group = QButtonGroup(self)
        self.leave_det_susp_yes = QRadioButton("Yes")
        self.leave_det_susp_no = QRadioButton("No")
        self.leave_det_susp_no.setChecked(True)
        for rb in [self.leave_det_susp_yes, self.leave_det_susp_no]:
            rb.setStyleSheet("font-size: 15px; color: #374151;")
            self.leave_det_susp_group.addButton(rb)
            susp_row.addWidget(rb)
        susp_row.addStretch()
        self.leave_det_susp_yes.toggled.connect(self._on_leave_det_susp_toggled)
        self.leave_det_susp_no.toggled.connect(self._update_leave_details_text)
        right_layout.addLayout(susp_row)

        self.leave_det_susp_details = QLineEdit()
        self.leave_det_susp_details.setPlaceholderText("Reason for suspension...")
        self.leave_det_susp_details.setStyleSheet("font-size: 15px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.leave_det_susp_details.setVisible(False)
        self.leave_det_susp_details.textChanged.connect(self._update_leave_details_text)
        right_layout.addWidget(self.leave_det_susp_details)

        right_layout.addStretch()
        split.addWidget(right_panel)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    def _on_leave_details_taken_toggled(self, checked):
        """Show/hide leave types when leave taken toggled."""
        self.leave_det_types_frame.setVisible(checked)
        self._update_leave_details_text()

    def _on_leave_det_type_toggled(self, type_key, checked):
        """Show/hide sliders when leave type toggled."""
        widgets = self.leave_det_type_widgets[type_key]
        widgets["sliders_widget"].setVisible(checked)
        self._update_leave_details_text()

    def _on_leave_det_slider_changed(self, type_key):
        """Update slider label and text."""
        widgets = self.leave_det_type_widgets[type_key]
        freq = widgets["freq_slider"].value()
        dur = widgets["dur_slider"].value()

        freq_labels = ["1x/week", "2x/week", "3x/week", "4x/week", "5x/week", "6x/week", "Daily"]
        widgets["freq_val"].setText(freq_labels[freq - 1])

        dur_mins = dur * 30
        if dur_mins >= 60:
            hours = dur_mins // 60
            mins = dur_mins % 60
            if mins > 0:
                widgets["dur_val"].setText(f"{hours}h {mins}m")
            else:
                widgets["dur_val"].setText(f"{hours} hours")
        else:
            widgets["dur_val"].setText(f"{dur_mins} mins")

        self._update_leave_details_text()

    def _on_leave_det_susp_toggled(self, checked):
        """Show/hide suspension details."""
        self.leave_det_susp_details.setVisible(checked)
        self._update_leave_details_text()

    def _update_leave_details_text(self):
        """Update leave details text based on selections."""
        parts = []

        if self.leave_det_taken_no.isChecked():
            parts.append("No leave taken in the last 12 months.")
        elif self.leave_det_taken_yes.isChecked():
            leave_parts = []
            for type_key, widgets in self.leave_det_type_widgets.items():
                if widgets["checkbox"].isChecked():
                    freq_labels = ["once weekly", "twice weekly", "3x weekly", "4x weekly", "5x weekly", "6x weekly", "daily"]
                    freq = freq_labels[widgets["freq_slider"].value() - 1]
                    dur = widgets["dur_slider"].value() * 30
                    if dur >= 60:
                        hours = dur // 60
                        mins = dur % 60
                        dur_str = f"{hours}h {mins}m" if mins else f"{hours} hours"
                    else:
                        dur_str = f"{dur} minutes"
                    type_name = widgets["checkbox"].text()
                    leave_parts.append(f"{type_name} leave ({freq}, {dur_str})")

            if leave_parts:
                parts.append("Leave taken: " + "; ".join(leave_parts) + ".")
            else:
                parts.append("Leave has been taken (types to be specified).")

        # Attendance
        attend_parts = []
        for att_key, cb in self.leave_det_attend.items():
            if cb.isChecked():
                attend_parts.append(cb.text().lower())
        if attend_parts:
            parts.append("Leave to attend: " + ", ".join(attend_parts) + ".")

        # Suspensions
        if self.leave_det_susp_yes.isChecked():
            reason = self.leave_det_susp_details.text().strip()
            if reason:
                parts.append(f"Leave has been suspended: {reason}.")
            else:
                parts.append("Leave has been suspended.")
        elif self.leave_det_susp_no.isChecked() and self.leave_det_taken_yes.isChecked():
            parts.append("Leave has not been suspended.")

        new_generated = " ".join(parts)
        if new_generated:
            self._update_text_preserving_additions(self.leave_details_text, new_generated, "leave_details")

    # ================================================================
    # SECTION 15: Additional Comments
    # ================================================================
    def _build_section_15_additional_comments(self):
        frame = self._create_section_frame("Additional Comments", "15", "#6b7280")
        layout = frame.layout()

        prompt_lbl = QLabel("Please consider the following:")
        prompt_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        layout.addWidget(prompt_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.additional_comments_text = self._create_text_edit("", 120)
        split.addWidget(self.additional_comments_text, 3)

        prompts_frame, self.additional_checkboxes = self._create_prompts_frame([
            "Please detail any other information or views you consider to be pertinent to the Annual Statutory Report"
        ], "#6b7280")
        prompts_frame.setFixedWidth(380)
        split.addWidget(prompts_frame)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    # ================================================================
    # SECTION 16: Unfit to Plead (if applicable)
    # ================================================================
    def _build_section_16_unfit_to_plead(self):
        frame = self._create_section_frame("Unfit to Plead (s37/41 after s24 DVCVA 2004 only)", "16", "#6b7280")
        layout = frame.layout()

        info_lbl = QLabel("For patients whose s37/41 order was made after a finding of unfit to plead (under s24 of the Domestic Violence, Crime and Victims Act 2004) only:")
        info_lbl.setWordWrap(True)
        info_lbl.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        layout.addWidget(info_lbl)

        split = QHBoxLayout()
        split.setSpacing(16)

        self.unfit_to_plead_text = self._create_text_edit("", 100)
        split.addWidget(self.unfit_to_plead_text, 3)

        prompts_frame, self.unfit_checkboxes = self._create_prompts_frame([
            "Do you consider that the patient is now fit to plead at Court for the offence which led to the current Order"
        ], "#6b7280")
        prompts_frame.setFixedWidth(380)
        split.addWidget(prompts_frame)

        layout.addLayout(split)
        self.form_layout.addWidget(frame)

    # ================================================================
    # Signature Section
    # ================================================================
    def _build_signature(self):
        frame = self._create_section_frame("Signature", "", "#374151")
        layout = frame.layout()

        row = QHBoxLayout()
        row.setSpacing(20)

        sig_lbl = QLabel("RC's signature:")
        sig_lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        row.addWidget(sig_lbl)

        self.signature_line = self._create_line_edit("Signed electronically")
        self.signature_line.setFixedWidth(300)
        row.addWidget(self.signature_line)

        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        row.addWidget(date_lbl)

        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        row.addWidget(self.sig_date)

        row.addStretch()
        layout.addLayout(row)

        info_lbl = QLabel("Please send the completed report to: mhcsmailbox@justice.gov.uk")
        info_lbl.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        layout.addWidget(info_lbl)

        self.form_layout.addWidget(frame)

    # ================================================================
    # Actions
    # ================================================================
    def _clear_form(self):
        reply = QMessageBox.question(
            self, "Clear Form", "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            # Clear all card editors
            for card in self.cards.values():
                card.editor.clear()

            # Destroy all popups and remove from stack
            for key, popup in list(self.popups.items()):
                if hasattr(self, 'popup_stack'):
                    self.popup_stack.removeWidget(popup)
                popup.deleteLater()
            self.popups.clear()

            # Clear raw notes and cached data
            self._extracted_raw_notes = []
            if hasattr(self, '_imported_report_data'):
                self._imported_report_data = {}
            if hasattr(self, '_imported_report_sections'):
                self._imported_report_sections = {}
            for attr in ['_progress_all_categorized', '_capacity_all_categorized', '_patient_attitude_all_categorized',
                         '_abscond_all_categorized', '_mappa_all_categorized', '_leave_report_all_categorized',
                         '_risk_addressed_all_categorized', '_behaviour_all_categorized']:
                if hasattr(self, attr):
                    setattr(self, attr, [])

            # Restore my details fields
            self._prefill()

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export MOJ ASR Form",
            f"MOJ_ASR_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            template_path = resource_path('templates', 'MOJ_ASR_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", f"MOJ ASR template not found at:\n{template_path}")
                return

            doc = Document(template_path)

            # Helper to get card content
            def get_card_text(key):
                if key in self.cards:
                    return self.cards[key].editor.toPlainText()
                return ""

            # Helper to fill table cells
            def fill_table_cell(table_idx, row_idx, col_idx, text):
                try:
                    cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        para.clear()
                        run = para.add_run(text)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                except Exception as e:
                    print(f"Error filling cell [{table_idx}][{row_idx}][{col_idx}]: {e}")

            # Helper to insert content after a section header
            def insert_after_header(header_text, content, para_offset=1):
                """Find header and insert content in the next empty paragraph."""
                if not content:
                    return
                for i, para in enumerate(doc.paragraphs):
                    if header_text.lower() in para.text.lower():
                        # Find the next empty or near-empty paragraph to insert content
                        target_idx = i + para_offset
                        if target_idx < len(doc.paragraphs):
                            target_para = doc.paragraphs[target_idx]
                            target_para.clear()
                            run = target_para.add_run(content)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.size = Pt(11)
                            # Set proper formatting: left indent, no hanging indent, left-aligned
                            target_para.paragraph_format.left_indent = Inches(0.24)
                            target_para.paragraph_format.first_line_indent = Inches(0)
                            target_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        return
                print(f"Header not found: {header_text}")

            # Get data from popup widgets
            patient_name = getattr(self, 'popup_patient_name', None)
            patient_name = patient_name.text() if patient_name else ""
            patient_dob = getattr(self, 'popup_dob', None)
            patient_dob = patient_dob.date().toString('dd/MM/yyyy') if patient_dob else ""
            hospital = getattr(self, 'popup_hospital', None)
            hospital = hospital.text() if hospital else ""
            mha_section = getattr(self, 'popup_mha_section', None)
            mha_section = mha_section.currentText() if mha_section else ""
            mhcs_ref = getattr(self, 'popup_mhcs_ref', None)
            mhcs_ref = mhcs_ref.text() if mhcs_ref else ""
            mha_section_date = getattr(self, 'popup_mha_section_date', None)
            mha_section_date = mha_section_date.date().toString('dd/MM/yyyy') if mha_section_date else ""
            other_detention = getattr(self, 'popup_other_detention', None)
            other_detention = other_detention.text() if other_detention else ""

            rc_name = getattr(self, 'popup_rc_name', None)
            rc_name = rc_name.text() if rc_name else ""
            rc_email = getattr(self, 'popup_rc_email', None)
            rc_email = rc_email.text() if rc_email else ""
            rc_job_title = getattr(self, 'popup_rc_job_title', None)
            rc_job_title = rc_job_title.text() if rc_job_title else ""
            rc_phone = getattr(self, 'popup_rc_phone', None)
            rc_phone = rc_phone.text() if rc_phone else ""
            mha_office_email = getattr(self, 'popup_mha_office_email', None)
            mha_office_email = mha_office_email.text() if mha_office_email else ""

            sig_date = getattr(self, 'popup_sig_date', None)
            sig_date = sig_date.date().toString('dd/MM/yyyy') if sig_date else ""

            # Fill Patient Details (Table 1)
            fill_table_cell(1, 1, 1, patient_name)
            fill_table_cell(1, 2, 1, patient_dob)
            fill_table_cell(1, 3, 1, hospital)
            fill_table_cell(1, 4, 1, mhcs_ref)
            fill_table_cell(1, 5, 1, mha_section)
            fill_table_cell(1, 6, 1, mha_section_date)
            fill_table_cell(1, 7, 1, other_detention)

            # Fill RC Details (Table 2)
            fill_table_cell(2, 1, 1, rc_name)
            fill_table_cell(2, 2, 1, rc_job_title)
            fill_table_cell(2, 3, 1, rc_phone)
            fill_table_cell(2, 4, 1, rc_email)
            fill_table_cell(2, 5, 1, mha_office_email)

            # Map section headers to card keys and paragraph offsets
            # The offset tells how many paragraphs after the header to insert content
            # Using shorter search strings to avoid apostrophe encoding issues
            section_map = [
                ("mental disorder, including the reasons", "mental_disorder", 1),
                ("attitude and behaviour in the last 12", "attitude_behaviour", 1),
                ("effect these have had on the patient", "addressing_issues", 1),
                ("attitude in the last 12 months:", "patient_attitude", 1),
                ("Capacity Issues:", "capacity", 1),
                ("Progress:", "progress", 1),
                ("key risk factors", "managing_risk", 1),
                ("risks have been addressed", "risk_addressed", 1),
                ("Abscond / Escape:", "abscond", 1),
                ("MAPPA coordinator", "mappa", 1),
                ("Victim Liaison Officer (VLO)", "victims", 1),
                ("consider the following:", "additional_comments", 1),
                ("unfit to plead", "unfit_to_plead", 1),
            ]

            # Insert content for each section
            for header, card_key, offset in section_map:
                content = get_card_text(card_key)
                insert_after_header(header, content, offset)

            # Handle leave report specially - insert after community leave checkbox table
            leave_content = get_card_text("leave_report")
            if leave_content:
                # Find paragraph after community leave table
                for i, para in enumerate(doc.paragraphs):
                    if "community leave taken" in para.text.lower():
                        # Insert content a few paragraphs after (after the checkbox table)
                        target_idx = i + 4
                        if target_idx < len(doc.paragraphs):
                            target_para = doc.paragraphs[target_idx]
                            target_para.clear()
                            run = target_para.add_run(leave_content)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.size = Pt(11)
                            # Set proper formatting: left indent, no hanging indent, left-aligned
                            target_para.paragraph_format.left_indent = Inches(0.24)
                            target_para.paragraph_format.first_line_indent = Inches(0)
                            target_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        break

            # Tick community leave checkboxes based on leave content
            # Table 3 structure: 5 rows x 5 cols
            # Row 0: Compassionate - col 0 (day checkbox), col 2 (overnight checkbox)
            # Row 1: Medical - col 0 (day checkbox), col 2 (overnight checkbox)
            # Row 2: Escorted - col 0 (day checkbox), col 2 (overnight checkbox)
            # Row 3: Unescorted - col 0 (day checkbox), col 2 (overnight checkbox)
            # Row 4: Long Term Escorted - col 0 (checkbox)
            leave_lower = leave_content.lower() if leave_content else ""
            has_escorted = "escorted leave:" in leave_lower
            has_unescorted = "unescorted leave:" in leave_lower

            def tick_cell(table, row, col):
                """Helper to tick a checkbox cell."""
                try:
                    cell = table.rows[row].cells[col]
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run("X")
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.bold = True
                except Exception as e:
                    print(f"Error ticking cell [{row}][{col}]: {e}")

            try:
                leave_table = doc.tables[3]

                # Row 2: Escorted (day col 0, overnight col 2)
                if has_escorted:
                    tick_cell(leave_table, 2, 0)  # Escorted day

                # Row 3: Unescorted (day col 0, overnight col 2)
                if has_unescorted:
                    tick_cell(leave_table, 3, 0)  # Unescorted day

                # Row 4: Long Term Escorted - tick if escorted content present
                if has_escorted:
                    tick_cell(leave_table, 4, 0)  # Long Term Escorted
            except Exception as e:
                print(f"Error ticking leave checkboxes: {e}")

            # Fill signature date (Table 4, Row 0, Column 3)
            try:
                if sig_date:
                    fill_table_cell(4, 0, 3, sig_date)
                    # Also fill RC name for signature (Table 4, Row 0, Column 1 - replaces "Please insert image")
                    fill_table_cell(4, 0, 1, rc_name)
            except Exception as e:
                print(f"Error filling signature: {e}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"MOJ ASR form exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        """Get form state for saving - card-based architecture."""
        # Save all card contents
        card_contents = {}
        for key, card in self.cards.items():
            card_contents[key] = card.editor.toPlainText()

        # Get gender from popup
        gender = "neutral"
        if hasattr(self, 'popup_gender_male') and self.popup_gender_male.isChecked():
            gender = "male"
        elif hasattr(self, 'popup_gender_female') and self.popup_gender_female.isChecked():
            gender = "female"

        # Get key popup values
        popup_data = {}
        if hasattr(self, 'popup_patient_name'):
            popup_data["patient_name"] = self.popup_patient_name.text()
        if hasattr(self, 'popup_dob'):
            popup_data["patient_dob"] = self.popup_dob.date().toString("yyyy-MM-dd")
        if hasattr(self, 'popup_nhs'):
            popup_data["nhs_number"] = self.popup_nhs.text()
        if hasattr(self, 'popup_hospital'):
            popup_data["hospital"] = self.popup_hospital.text()
        if hasattr(self, 'popup_mha_section'):
            popup_data["mha_section"] = self.popup_mha_section.currentText()
        if hasattr(self, 'popup_mhcs_ref'):
            popup_data["mhcs_ref"] = self.popup_mhcs_ref.text()
        if hasattr(self, 'popup_mha_section_date'):
            popup_data["mha_section_date"] = self.popup_mha_section_date.date().toString("yyyy-MM-dd")
        if hasattr(self, 'popup_other_detention'):
            popup_data["other_detention"] = self.popup_other_detention.text()
        if hasattr(self, 'popup_rc_name'):
            popup_data["rc_name"] = self.popup_rc_name.text()
        if hasattr(self, 'popup_rc_job_title'):
            popup_data["rc_job_title"] = self.popup_rc_job_title.text()
        if hasattr(self, 'popup_rc_phone'):
            popup_data["rc_phone"] = self.popup_rc_phone.text()
        if hasattr(self, 'popup_rc_email'):
            popup_data["rc_email"] = self.popup_rc_email.text()
        if hasattr(self, 'popup_mha_office_email'):
            popup_data["mha_office_email"] = self.popup_mha_office_email.text()
        if hasattr(self, 'popup_sig_date'):
            popup_data["sig_date"] = self.popup_sig_date.date().toString("yyyy-MM-dd")

        return {
            "card_contents": card_contents,
            "gender": gender,
            "popup_data": popup_data,
        }

    def load_state(self, state: dict):
        """Load form state - card-based architecture."""
        if not state:
            return

        # Load card contents
        card_contents = state.get("card_contents", {})
        for key, content in card_contents.items():
            if key in self.cards:
                self.cards[key].editor.setPlainText(content)

        # Load gender
        g = state.get("gender", "neutral")
        if g == "male" and hasattr(self, 'popup_gender_male'):
            self.popup_gender_male.setChecked(True)
        elif g == "female" and hasattr(self, 'popup_gender_female'):
            self.popup_gender_female.setChecked(True)

        # Load popup data
        popup_data = state.get("popup_data", {})
        if hasattr(self, 'popup_patient_name') and popup_data.get("patient_name"):
            self.popup_patient_name.setText(popup_data["patient_name"])
        if hasattr(self, 'popup_dob') and popup_data.get("patient_dob"):
            self.popup_dob.setDate(QDate.fromString(popup_data["patient_dob"], "yyyy-MM-dd"))
        if hasattr(self, 'popup_nhs') and popup_data.get("nhs_number"):
            self.popup_nhs.setText(popup_data["nhs_number"])
        if hasattr(self, 'popup_hospital') and popup_data.get("hospital"):
            self.popup_hospital.setText(popup_data["hospital"])
        if hasattr(self, 'popup_mha_section') and popup_data.get("mha_section"):
            idx = self.popup_mha_section.findText(popup_data["mha_section"])
            if idx >= 0:
                self.popup_mha_section.setCurrentIndex(idx)
        if hasattr(self, 'popup_mhcs_ref') and popup_data.get("mhcs_ref"):
            self.popup_mhcs_ref.setText(popup_data["mhcs_ref"])
        if hasattr(self, 'popup_mha_section_date') and popup_data.get("mha_section_date"):
            self.popup_mha_section_date.setDate(QDate.fromString(popup_data["mha_section_date"], "yyyy-MM-dd"))
        if hasattr(self, 'popup_other_detention') and popup_data.get("other_detention"):
            self.popup_other_detention.setText(popup_data["other_detention"])
        if hasattr(self, 'popup_rc_name') and popup_data.get("rc_name"):
            self.popup_rc_name.setText(popup_data["rc_name"])
        if hasattr(self, 'popup_rc_job_title') and popup_data.get("rc_job_title"):
            self.popup_rc_job_title.setText(popup_data["rc_job_title"])
        if hasattr(self, 'popup_rc_phone') and popup_data.get("rc_phone"):
            self.popup_rc_phone.setText(popup_data["rc_phone"])
        if hasattr(self, 'popup_rc_email') and popup_data.get("rc_email"):
            self.popup_rc_email.setText(popup_data["rc_email"])
        if hasattr(self, 'popup_mha_office_email') and popup_data.get("mha_office_email"):
            self.popup_mha_office_email.setText(popup_data["mha_office_email"])
        if hasattr(self, 'popup_sig_date') and popup_data.get("sig_date"):
            self.popup_sig_date.setDate(QDate.fromString(popup_data["sig_date"], "yyyy-MM-dd"))
