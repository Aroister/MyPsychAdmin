# ================================================================
#  CTO3 FORM PAGE — Notice of Recall to Hospital
#  Mental Health Act 1983 - Form CTO3 Regulation 6(3)(a)
#  Section 17E — Community treatment order: notice of recall
#  CARD/POPUP LAYOUT with ResizableSection
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QRadioButton, QButtonGroup, QComboBox, QSpinBox, QCompleter,
    QStyleFactory, QSlider, QStackedWidget, QSplitter, QSizePolicy
)
from utils.resource_path import resource_path


# ================================================================
# RESIZABLE SECTION WITH DRAG BAR
# ================================================================
class ResizableSection(QFrame):
    """Section with a draggable bottom edge to resize height."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._min_height = 120
        self._max_height = 600
        self._content = None
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0

        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(0)

        # Content container
        self._content_container = QWidget()
        self._content_layout = QVBoxLayout(self._content_container)
        self._content_layout.setContentsMargins(12, 8, 12, 8)
        self._layout.addWidget(self._content_container, 1)

        # Drag handle at bottom
        self._drag_handle = QFrame()
        self._drag_handle.setFixedHeight(8)
        self._drag_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        self._drag_handle.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
                border-radius: 2px;
            }
            QFrame:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.3 #dc2626, stop:0.7 #dc2626, stop:1 transparent);
            }
        """)
        self._drag_handle.mousePressEvent = self._handle_press
        self._drag_handle.mouseMoveEvent = self._handle_move
        self._drag_handle.mouseReleaseEvent = self._handle_release
        self._layout.addWidget(self._drag_handle)

    def set_content(self, widget):
        self._content = widget
        self._content_layout.addWidget(widget)

    def set_content_height(self, h):
        h = max(self._min_height, min(self._max_height, h))
        self._content_container.setFixedHeight(h)

    def _handle_press(self, event):
        self._dragging = True
        self._drag_start_y = event.globalPosition().y()
        self._drag_start_height = self._content_container.height()

    def _handle_move(self, event):
        if self._dragging:
            delta = event.globalPosition().y() - self._drag_start_y
            new_height = self._drag_start_height + delta
            new_height = max(self._min_height, min(self._max_height, new_height))
            self._content_container.setFixedHeight(int(new_height))

    def _handle_release(self, event):
        self._dragging = False


# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []

from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# NO-WHEEL SLIDER
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# NO-WHEEL COMBOBOX
# ================================================================
class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelSpinBox(QSpinBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelTimeEdit(QTimeEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# CTO3 CARD WIDGET
# ================================================================
class CTO3CardWidget(QFrame):
    """Clickable card with editable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("cto3Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)

        self.setStyleSheet("""
            QFrame#cto3Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#cto3Card:hover {
                border-color: #dc2626;
                background: #fef2f2;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(6)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 19px; font-weight: 700; color: #dc2626;")
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 19px;
                color: #374151;
                background: transparent;
                border: none;
                padding: 0;
            }
        """)
        self.content.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        layout.addWidget(self.content, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=17)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_active(self, active: bool):
        """Set active state - shows persistent highlight."""
        if active:
            self.setStyleSheet("""
                QFrame#cto3Card {
                    background: #fef2f2;
                    border: 2px solid #dc2626;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#cto3Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#cto3Card:hover {
                    border-color: #dc2626;
                    background: #fef2f2;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# CTO3 TOOLBAR
# ================================================================
class CTO3Toolbar(QWidget):
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            CTO3Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN CTO3 FORM PAGE - Card/Popup Layout
# ================================================================
class CTO3FormPage(QWidget):
    """Page for completing MHA Form CTO3 - Notice of Recall to Hospital."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race", "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}

        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {"full_name": details[1] or "", "email": details[7] or ""}

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(50)
        header.setStyleSheet("background: #dc2626; border-bottom: 1px solid #b91c1c;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 5px;
                font-size: 18px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO3 — Notice of Recall to Hospital")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - cards on left, popup on right with splitter
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(12, 12, 12, 12)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #dc2626, stop:0.7 #dc2626, stop:1 transparent);
            }
        """)

        # Left: Cards column (scrollable)
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(0, 0, 8, 0)
        self.cards_layout.setSpacing(8)

        # Create cards
        self._create_details_card()
        self._create_grounds_card()
        self._create_signatures_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        self.main_splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 10px;
            }
        """)

        # Create popup panels
        self._create_details_popup()
        self._create_grounds_popup()
        self._create_signatures_popup()

        # Initialize cards with default date values
        self._update_signatures_card()

        self.main_splitter.addWidget(self.popup_stack)
        self.main_splitter.setSizes([280, 600])
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        # Connect live preview updates
        self._connect_grounds_live_preview()

        # Show first popup by default
        self._on_card_clicked("details")

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"details": 0, "grounds": 1, "signatures": 2}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            for k, card in self.cards.items():
                card.set_active(k == key)

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_details_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 350
        card = CTO3CardWidget("Details", "details")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["details"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_grounds_card(self):
        section = ResizableSection()
        section.set_content_height(200)
        section._min_height = 120
        section._max_height = 400
        card = CTO3CardWidget("Reason for Recall", "grounds")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["grounds"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signatures_card(self):
        section = ResizableSection()
        section.set_content_height(150)
        section._min_height = 100
        section._max_height = 300
        card = CTO3CardWidget("Signature", "signatures")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signatures"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_details_popup(self):
        """Combined popup for patient, hospital, RC details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        # Header
        header = QLabel("Patient, Hospital & RC Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #dc2626;")
        popup_layout.addWidget(header)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        form = QWidget()
        form.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(form)
        form_layout.setContentsMargins(0, 0, 8, 0)
        form_layout.setSpacing(10)

        # Patient section
        patient_header = QLabel("Patient")
        patient_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 4px;")
        form_layout.addWidget(patient_header)

        self.patient_name = self._create_line_edit("Patient full name")
        self.patient_name.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.patient_name)

        self.patient_address = self._create_line_edit("Patient address")
        self.patient_address.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.patient_address)

        # Demographics row (age, gender, ethnicity - visible for clinical reasons)
        demo_row = QHBoxLayout()
        demo_row.setSpacing(8)

        # Age label and spin
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        demo_row.addWidget(age_lbl)
        self.age_spin = NoWheelSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setFixedWidth(60)
        self.age_spin.setStyleSheet("QSpinBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 18px; }")
        demo_row.addWidget(self.age_spin)

        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("O")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 17px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
        self.gender_group.addButton(self.gender_male, 0)
        self.gender_group.addButton(self.gender_female, 1)
        self.gender_group.addButton(self.gender_other, 2)
        demo_row.addWidget(self.gender_male)
        demo_row.addWidget(self.gender_female)
        demo_row.addWidget(self.gender_other)

        # Ethnicity combo - visible for clinical reasons
        self.ethnicity_combo = NoWheelComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(130)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 18px; }")
        demo_row.addWidget(self.ethnicity_combo)
        demo_row.addStretch()
        form_layout.addLayout(demo_row)

        # Hospital section
        hosp_header = QLabel("Hospital")
        hosp_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(hosp_header)

        self.hospital_name = self._create_line_edit("Hospital name")
        self.hospital_name.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.hospital_name)

        self.hospital_address = self._create_line_edit("Hospital address")
        self.hospital_address.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.hospital_address)

        # RC section
        rc_header = QLabel("Responsible Clinician")
        rc_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(rc_header)

        self.rc_name = self._create_line_edit("RC full name")
        self.rc_name.textChanged.connect(self._update_details_card)
        self.rc_name.textChanged.connect(self._update_signatures_card)
        form_layout.addWidget(self.rc_name)

        form_layout.addStretch()
        scroll.setWidget(form)
        popup_layout.addWidget(scroll, 1)

        self.popup_stack.addWidget(popup)

    def _create_grounds_popup(self):
        """Popup for grounds with live preview."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(0, 0, 0, 0)
        popup_layout.setSpacing(0)

        # Hidden label for state storage (preview removed - card auto-syncs)
        self.grounds_preview = QLabel("")
        self.grounds_preview.hide()

        # Scrollable form below
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("QScrollArea { background: white; border: none; }")

        form_container = QWidget()
        form_container.setStyleSheet("background: white;")
        form_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(12, 8, 12, 12)
        form_layout.setSpacing(8)

        # ===== REASON FOR RECALL SECTION (at TOP) =====
        rr_frame = QFrame()
        rr_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        rr_frame.setStyleSheet("QFrame { background: #fef2f2; border: none; border-radius: 6px; }")
        rr_layout = QVBoxLayout(rr_frame)
        rr_layout.setContentsMargins(10, 8, 10, 8)
        rr_layout.setSpacing(4)
        rr_header = QLabel("Reason for Recall")
        rr_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #991b1b;")
        rr_layout.addWidget(rr_header)

        # Radio button group for (a) vs (b) - mutually exclusive
        self.reason_group = QButtonGroup(self)
        self.reason_group.setExclusive(True)

        # (a) Treatment required AND risk
        self.reason_a = QRadioButton("(a) Treatment required AND risk")
        self.reason_a.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.reason_a.toggled.connect(self._on_reason_a_toggled)
        self.reason_group.addButton(self.reason_a)
        rr_layout.addWidget(self.reason_a)

        # (b) Failed to comply with condition
        self.reason_b = QRadioButton("(b) Failure to comply with condition of CTO to make yourself available for")
        self.reason_b.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.reason_b.toggled.connect(self._on_reason_b_toggled)
        self.reason_group.addButton(self.reason_b)
        rr_layout.addWidget(self.reason_b)

        # (b) sub-options: (i) and (ii) as radio buttons
        self.reason_b_options = QWidget()
        reason_b_layout = QVBoxLayout(self.reason_b_options)
        reason_b_layout.setContentsMargins(16, 2, 0, 2)
        reason_b_layout.setSpacing(2)

        self.condition_group = QButtonGroup(self)
        self.condition_group.setExclusive(True)

        self.condition_i_rb = QRadioButton("(i) consideration of extension of the community treatment period under section 20A")
        self.condition_i_rb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.condition_i_rb.toggled.connect(self._update_grounds_preview)
        self.condition_group.addButton(self.condition_i_rb)
        reason_b_layout.addWidget(self.condition_i_rb)
        self.condition_ii_rb = QRadioButton("(ii) enabling a Part 4A certificate to be given")
        self.condition_ii_rb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.condition_ii_rb.toggled.connect(self._update_grounds_preview)
        self.condition_group.addButton(self.condition_ii_rb)
        reason_b_layout.addWidget(self.condition_ii_rb)
        self.reason_b_options.hide()
        rr_layout.addWidget(self.reason_b_options)

        form_layout.addWidget(rr_frame)

        # ===== MENTAL DISORDER SECTION (shown when (a) selected) =====
        self.md_frame = QFrame()
        self.md_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        self.md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 6px; }")
        md_layout = QVBoxLayout(self.md_frame)
        md_layout.setContentsMargins(10, 8, 10, 8)
        md_layout.setSpacing(4)
        md_header = QLabel("Mental Disorder (ICD-10)")
        md_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        # Primary diagnosis dropdown with grouped items
        self.dx_primary = NoWheelComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select primary diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_primary.currentTextChanged.connect(self._update_grounds_preview)
        md_layout.addWidget(self.dx_primary)

        # Secondary diagnosis dropdown
        self.dx_secondary = NoWheelComboBox()
        self.dx_secondary.setEditable(True)
        self.dx_secondary.lineEdit().setPlaceholderText("Secondary diagnosis (optional)...")
        self.dx_secondary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_secondary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_secondary.addItem(f"── {group_name} ──", None)
            idx = self.dx_secondary.count() - 1
            self.dx_secondary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_secondary.addItem(dx, dx)
        completer2 = QCompleter(ICD10_FLAT, self.dx_secondary)
        completer2.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer2.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_secondary.setCompleter(completer2)
        self.dx_secondary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_secondary.currentTextChanged.connect(self._update_grounds_preview)
        md_layout.addWidget(self.dx_secondary)

        self.md_frame.hide()  # Hidden by default, shown when (a) selected
        form_layout.addWidget(self.md_frame)

        # ===== LEGAL CRITERIA SECTION (shown when (a) selected) =====
        self.lc_frame = QFrame()
        self.lc_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        self.lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 6px; }")
        lc_layout = QVBoxLayout(self.lc_frame)
        lc_layout.setContentsMargins(10, 8, 10, 8)
        lc_layout.setSpacing(4)
        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)
        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._update_grounds_preview)
        nature_opt_layout.addWidget(self.relapsing_cb)
        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._update_grounds_preview)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)
        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._update_grounds_preview)
        nature_opt_layout.addWidget(self.chronic_cb)
        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)
        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(100)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)
        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)
        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._update_grounds_preview)
        degree_opt_layout.addWidget(self.degree_details)
        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity section (A4-style)
        nec_lbl = QLabel("Necessity (Risk if not recalled):")
        nec_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health with sub-options
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)
        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._update_grounds_preview)
        mh_opt_layout.addWidget(self.poor_compliance_cb)
        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._update_grounds_preview)
        mh_opt_layout.addWidget(self.limited_insight_cb)
        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._update_grounds_preview)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety with Self/Others sub-options
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # Self section
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)
        self.self_neglect_cb = QCheckBox("Self neglect")
        self.self_neglect_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_neglect_cb.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_neglect_cb)
        self.self_risky_cb = QCheckBox("Risky situations")
        self.self_risky_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_risky_cb.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_risky_cb)
        self.self_harm_detail_cb = QCheckBox("Self harm")
        self.self_harm_detail_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.self_harm_detail_cb.toggled.connect(self._update_grounds_preview)
        self_opt_layout.addWidget(self.self_harm_detail_cb)
        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others section
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)
        self.others_violence_cb = QCheckBox("Violence to others")
        self.others_violence_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_violence_cb.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_violence_cb)
        self.others_verbal_cb = QCheckBox("Verbal aggression")
        self.others_verbal_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_verbal_cb.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_verbal_cb)
        self.others_sexual_cb = QCheckBox("Sexual violence")
        self.others_sexual_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.others_sexual_cb.toggled.connect(self._update_grounds_preview)
        others_opt_layout.addWidget(self.others_sexual_cb)
        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        self.lc_frame.hide()  # Hidden by default, shown when (a) selected
        form_layout.addWidget(self.lc_frame)

        form_layout.addStretch()
        form_scroll.setWidget(form_container)
        popup_layout.addWidget(form_scroll, 1)

        self.popup_stack.addWidget(popup)

    def _create_signatures_popup(self):
        """Popup for RC signature details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("RC Signature")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #7c3aed;")
        popup_layout.addWidget(header)

        # Form
        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        # Date row
        date_row = QHBoxLayout()
        date_row.setSpacing(8)
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        date_row.addWidget(date_lbl)
        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        self.sig_date.dateChanged.connect(self._update_signatures_card)
        date_row.addWidget(self.sig_date)
        date_row.addStretch()
        form_layout.addLayout(date_row)

        # Time row
        time_row = QHBoxLayout()
        time_row.setSpacing(8)
        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        time_row.addWidget(time_lbl)
        self.sig_time = self._create_time_edit()
        self.sig_time.setFixedWidth(100)
        self.sig_time.timeChanged.connect(self._update_signatures_card)
        time_row.addWidget(self.sig_time)
        time_row.addStretch()
        form_layout.addLayout(time_row)

        notice = QLabel("The notice must be served on the patient personally or delivered to the patient's usual or last known address.")
        notice.setWordWrap(True)
        notice.setStyleSheet("font-size: 17px; color: #6b7280; font-style: italic; padding: 8px; background: #f3f4f6; border-radius: 6px; margin-top: 12px;")
        form_layout.addWidget(notice)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_details_card(self):
        parts = []
        if self.patient_name.text().strip():
            parts.append(self.patient_name.text().strip())
        if self.patient_address.text().strip():
            parts.append(self.patient_address.text().strip()[:40] + "...")
        if self.hospital_name.text().strip():
            parts.append(self.hospital_name.text().strip())
        if self.hospital_address.text().strip():
            parts.append(self.hospital_address.text().strip()[:40] + "...")
        self.cards["details"].set_content_text("\n".join(parts) if parts else "Click to enter details")

    def _update_grounds_card(self):
        self.cards["grounds"].set_content_text(self.grounds_preview.text())

    def _update_signatures_card(self):
        parts = []
        if self.sig_date.date().isValid():
            parts.append(self.sig_date.date().toString('dd MMM yyyy'))
        if self.sig_time.time().isValid():
            parts.append(self.sig_time.time().toString('HH:mm'))
        self.cards["signatures"].set_content_text("\n".join(parts) if parts else "Click to enter details")

    # ----------------------------------------------------------------
    # HELPERS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("QLineEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 8px; font-size: 18px; } QLineEdit:focus { border-color: #dc2626; }")
        return edit

    def _create_date_edit(self) -> NoWheelDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 6px; font-size: 18px; }")
        return date_edit

    def _create_time_edit(self) -> NoWheelTimeEdit:
        time_edit = NoWheelTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("QTimeEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 6px; font-size: 18px; }")
        return time_edit

    # ----------------------------------------------------------------
    # TOGGLE HANDLERS
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_grounds_preview()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_grounds_preview()

    def _on_reason_a_toggled(self, checked):
        # Show/hide Mental Disorder and Legal Criteria sections when (a) is selected
        self.md_frame.setVisible(checked)
        self.lc_frame.setVisible(checked)
        if not checked:
            # Clear mental disorder selections
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            # Clear legal criteria selections
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            # Clear necessity selections
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_reason_b_toggled(self, checked):
        self.reason_b_options.setVisible(checked)
        if not checked:
            self.condition_group.setExclusive(False)
            self.condition_i_rb.setChecked(False)
            self.condition_ii_rb.setChecked(False)
            self.condition_group.setExclusive(True)
        self._update_grounds_preview()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_grounds_preview()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            self.self_neglect_cb.setChecked(False)
            self.self_risky_cb.setChecked(False)
            self.self_harm_detail_cb.setChecked(False)
        self._update_grounds_preview()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            self.others_violence_cb.setChecked(False)
            self.others_verbal_cb.setChecked(False)
            self.others_sexual_cb.setChecked(False)
        self._update_grounds_preview()

    # ----------------------------------------------------------------
    # LIVE PREVIEW
    # ----------------------------------------------------------------
    def _connect_grounds_live_preview(self):
        self.patient_name.textChanged.connect(self._update_grounds_preview)
        self.gender_male.toggled.connect(self._update_grounds_preview)
        self.gender_female.toggled.connect(self._update_grounds_preview)
        self.gender_other.toggled.connect(self._update_grounds_preview)

    def _update_grounds_preview(self):
        text = self._generate_grounds_text()
        self.grounds_preview.setText(text)
        self._update_grounds_card()

    def _generate_grounds_text(self) -> str:
        # CTO3 is addressed directly to the patient using "you"
        paragraphs = []

        # Para 1: Demographics + Diagnosis
        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        if self.dx_primary.currentText().strip():
            diagnoses.append(self.dx_primary.currentText().strip())
        if self.dx_secondary.currentText().strip():
            diagnoses.append(self.dx_secondary.currentText().strip())

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"You are a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"You suffer from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for you to receive medical treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for you to receive medical treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for you to receive medical treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        # Para 2: Reason for Recall - Treatment required with necessity/risk
        para2_parts = []
        if self.reason_a.isChecked():
            para2_parts.append("You require treatment for your mental disorder in hospital.")

            # Build necessity items
            necessity_items = []
            if self.health_cb.isChecked():
                necessity_items.append("your health")
            if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
                necessity_items.append("your own safety")
            if self.safety_cb.isChecked() and self.others_cb.isChecked():
                necessity_items.append("the safety of others")

            if necessity_items:
                if len(necessity_items) == 1:
                    para2_parts.append(f"If not recalled, there would be a risk of harm to {necessity_items[0]}.")
                elif len(necessity_items) == 2:
                    para2_parts.append(f"If not recalled, there would be a risk of harm to {necessity_items[0]} and {necessity_items[1]}.")
                else:
                    para2_parts.append(f"If not recalled, there would be a risk of harm to {necessity_items[0]}, {necessity_items[1]}, and {necessity_items[2]}.")

            # Health details
            if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
                mh_reasons = []
                if self.poor_compliance_cb.isChecked():
                    mh_reasons.append("non-compliance")
                if self.limited_insight_cb.isChecked():
                    mh_reasons.append("limited insight")
                if mh_reasons:
                    para2_parts.append(f"Your mental health would deteriorate due to {'/'.join(mh_reasons)}.")
                else:
                    para2_parts.append("Your mental health would deteriorate.")

            if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
                details = self.physical_health_details.text().strip()
                if details:
                    para2_parts.append(f"There are also concerns about your physical health: {details}.")
                else:
                    para2_parts.append("There are also concerns about your physical health.")

            # Safety - Self details
            if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
                self_risks = []
                if self.self_neglect_cb.isChecked():
                    self_risks.append("self neglect")
                if self.self_risky_cb.isChecked():
                    self_risks.append("placing yourself in risky situations")
                if self.self_harm_detail_cb.isChecked():
                    self_risks.append("self harm")
                if self_risks:
                    para2_parts.append(f"There is a risk to your safety through {', '.join(self_risks)}.")

            # Safety - Others details
            if self.safety_cb.isChecked() and self.others_cb.isChecked():
                others_risks = []
                if self.others_violence_cb.isChecked():
                    others_risks.append("violence to others")
                if self.others_verbal_cb.isChecked():
                    others_risks.append("verbal aggression")
                if self.others_sexual_cb.isChecked():
                    others_risks.append("sexual violence")
                if others_risks:
                    para2_parts.append(f"There is a risk to others through {', '.join(others_risks)}.")

        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        # Para 3: Failed to comply with condition
        para3_parts = []
        if self.reason_b.isChecked():
            condition = ""
            if self.condition_i_rb.isChecked():
                condition = "consideration of extension of the community treatment period under section 20A"
            elif self.condition_ii_rb.isChecked():
                condition = "enabling a Part 4A certificate to be given"

            if condition:
                para3_parts.append(f"You have failed to comply with the condition of the CTO to make yourself available for {condition}.")

        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        return "\n\n".join(paragraphs)

    # ----------------------------------------------------------------
    # EDITOR FOCUS TRACKING
    # ----------------------------------------------------------------
    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.patient_name.clear()
            self.patient_address.clear()
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.rc_name.clear()
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            # Clear reason radio buttons
            self.reason_group.setExclusive(False)
            self.reason_a.setChecked(False)
            self.reason_b.setChecked(False)
            self.reason_group.setExclusive(True)
            # Hide the conditional sections
            self.md_frame.hide()
            self.lc_frame.hide()
            self.reason_b_options.hide()
            # Clear necessity checkboxes
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            # Clear condition radio buttons
            self.condition_group.setExclusive(False)
            self.condition_i_rb.setChecked(False)
            self.condition_ii_rb.setChecked(False)
            self.condition_group.setExclusive(True)
            self.grounds_preview.setText("")
            self.sig_date.setDate(QDate.currentDate())
            self.sig_time.setTime(QTime.currentTime())
            for card in self.cards.values():
                card.set_content_text("")
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(self, "Export Form CTO3", f"Form_CTO3_{datetime.now().strftime('%Y%m%d')}.docx", "Word Documents (*.docx)")
        if not file_path:
            return
        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_CTO3_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO3 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)

            # Clean ALL paragraphs - remove permission markers and convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), 'FFFED5')
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), 'FFFED5')

            def set_entry_box(para, content=""):
                """Set entry box with bold gold brackets [content]"""
                if not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr = bracket_open._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'FFFED5')
                rPr2.append(shd2)
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr3 = bracket_close._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)

            def highlight_yellow(para):
                """Rebuild paragraph with cream highlight"""
                text = ''.join(run.text for run in para.runs)
                if not text:
                    text = para.text
                if not text.strip():
                    return
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                    rPr = para.runs[0]._element.get_or_add_rPr()
                    for old_shd in rPr.findall(qn('w:shd')):
                        rPr.remove(old_shd)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)
                else:
                    run = para.add_run(text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    rPr = run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Debug: Print paragraph indices
            print("DEBUG CTO3: Paragraphs 0-40:")
            for i in range(min(40, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)[:60]
                print(f"DEBUG: Para {i}: {txt}")

            # Patient name entry box with gold brackets (para 4)
            patient_text = self.patient_name.text().strip()
            if self.patient_address.text().strip():
                patient_text += ", " + self.patient_address.text().strip()
            set_entry_box(paragraphs[4], patient_text if patient_text else "")

            # Hospital entry box with gold brackets (para 6)
            hospital_text = self.hospital_name.text().strip()
            if self.hospital_address.text().strip():
                hospital_text += ", " + self.hospital_address.text().strip()
            set_entry_box(paragraphs[6], hospital_text if hospital_text else "")

            # Section (a) and (b) - user selects one, other gets strikethrough
            # Track paragraph indices for each section
            section_a_paras = []
            section_b_paras = []  # Main (b) paragraph only
            section_b_i_para = None  # (i) paragraph index
            section_b_ii_para = None  # (ii) paragraph index

            # Section (a) - Format the entire section with gold brackets and cream highlight
            # (a) [In my opinion,
            # (i) you require treatment in hospital for mental disorder,
            # AND
            # (ii) there would be a risk of harm...for that purpose.]
            # This opinion is founded on the following grounds—

            # Section (a) paragraphs based on actual template structure:
            # Para 9: "In my opinion,"
            # Para 10: "you require treatment in hospital..."
            # Para 11: "AND"
            # Para 12: "there would be a risk of harm..."
            # Para 13: "This opinion is founded..."
            # Para 14: entry box
            for i in range(5, min(25, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                txt_lower = txt.lower()

                # "In my opinion" - cream highlight with gold bracket at start (section a)
                if "in my opinion" in txt_lower and "failed" not in txt_lower:
                    print(f"DEBUG: Found 'In my opinion' at para {i}: {txt[:50]}")
                    section_a_paras.append(i)
                    para_a = paragraphs[i]
                    para_a_text = txt.strip()
                    for run in para_a.runs:
                        run.text = ""
                    while len(para_a.runs) > 1:
                        para_a._element.remove(para_a.runs[-1]._element)
                    if para_a.runs:
                        para_a.runs[0].text = '['
                        para_a.runs[0].font.bold = True
                        para_a.runs[0].font.color.rgb = BRACKET_COLOR
                    a_content = para_a.add_run(para_a_text)
                    a_content.font.name = 'Arial'
                    a_content.font.size = Pt(12)
                    rPr_ac = a_content._element.get_or_add_rPr()
                    shd_ac = OxmlElement('w:shd')
                    shd_ac.set(qn('w:val'), 'clear')
                    shd_ac.set(qn('w:color'), 'auto')
                    shd_ac.set(qn('w:fill'), 'FFFED5')
                    rPr_ac.append(shd_ac)

                # "you require treatment" - cream highlight (section a)
                elif "require treatment" in txt_lower:
                    print(f"DEBUG: Found 'require treatment' at para {i}: {txt[:50]}")
                    section_a_paras.append(i)
                    highlight_yellow(paragraphs[i])

                # AND - cream highlight (section a)
                elif txt.strip().upper() == "AND":
                    print(f"DEBUG: Found AND at para {i}")
                    section_a_paras.append(i)
                    highlight_yellow(paragraphs[i])

                # "there would be a risk of harm" - cream highlight with gold bracket at end (section a)
                elif "risk of harm" in txt_lower or ("risk" in txt_lower and "harm" in txt_lower and "safety" in txt_lower):
                    print(f"DEBUG: Found 'risk of harm' at para {i}: {txt[:50]}")
                    section_a_paras.append(i)
                    para_ii = paragraphs[i]
                    para_ii_text = txt.strip()
                    for run in para_ii.runs:
                        run.text = ""
                    while len(para_ii.runs) > 1:
                        para_ii._element.remove(para_ii.runs[-1]._element)
                    if para_ii.runs:
                        para_ii.runs[0].text = para_ii_text
                        para_ii.runs[0].font.name = 'Arial'
                        para_ii.runs[0].font.size = Pt(12)
                        rPr_ii = para_ii.runs[0]._element.get_or_add_rPr()
                        shd_ii = OxmlElement('w:shd')
                        shd_ii.set(qn('w:val'), 'clear')
                        shd_ii.set(qn('w:color'), 'auto')
                        shd_ii.set(qn('w:fill'), 'FFFED5')
                        rPr_ii.append(shd_ii)
                    ii_close = para_ii.add_run(']')
                    ii_close.font.bold = True
                    ii_close.font.color.rgb = BRACKET_COLOR

                # "This opinion is founded on the following grounds—" - cream highlight (section a)
                elif "this opinion is founded" in txt_lower:
                    print(f"DEBUG: Found 'This opinion is founded' at para {i}")
                    section_a_paras.append(i)
                    highlight_yellow(paragraphs[i])

            # Grounds entry box with gold brackets (para 14) - part of section (a)
            section_a_paras.append(14)  # Track for strikethrough when (b) is selected
            if self.reason_a.isChecked():
                grounds_text = self.cards["grounds"].get_content()
                set_entry_box(paragraphs[14], grounds_text if grounds_text.strip() else "")
            else:
                # When (b) is selected, leave the grounds box empty (will be struck through)
                set_entry_box(paragraphs[14], "")

            # Find and format "[If you need to continue...]" paragraphs
            for i in range(12, min(30, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "If you need to continue" in txt:
                    para_cont = paragraphs[i]
                    for run in para_cont.runs:
                        run.text = ""
                    while len(para_cont.runs) > 1:
                        para_cont._element.remove(para_cont.runs[-1]._element)
                    if para_cont.runs:
                        para_cont.runs[0].text = '[If you need to continue on a separate sheet please indicate here '
                        para_cont.runs[0].font.name = 'Arial'
                        para_cont.runs[0].font.size = Pt(12)
                    cb_open = para_cont.add_run('[')
                    cb_open.font.bold = True
                    cb_open.font.color.rgb = BRACKET_COLOR
                    rPr_cbo = cb_open._element.get_or_add_rPr()
                    shd_cbo = OxmlElement('w:shd')
                    shd_cbo.set(qn('w:val'), 'clear')
                    shd_cbo.set(qn('w:color'), 'auto')
                    shd_cbo.set(qn('w:fill'), 'FFFED5')
                    rPr_cbo.append(shd_cbo)
                    cb_space = para_cont.add_run(' ')
                    rPr_cbs = cb_space._element.get_or_add_rPr()
                    shd_cbs = OxmlElement('w:shd')
                    shd_cbs.set(qn('w:val'), 'clear')
                    shd_cbs.set(qn('w:color'), 'auto')
                    shd_cbs.set(qn('w:fill'), 'FFFED5')
                    rPr_cbs.append(shd_cbs)
                    cb_close = para_cont.add_run(']')
                    cb_close.font.bold = True
                    cb_close.font.color.rgb = BRACKET_COLOR
                    rPr_cbc = cb_close._element.get_or_add_rPr()
                    shd_cbc = OxmlElement('w:shd')
                    shd_cbc.set(qn('w:val'), 'clear')
                    shd_cbc.set(qn('w:color'), 'auto')
                    shd_cbc.set(qn('w:fill'), 'FFFED5')
                    rPr_cbc.append(shd_cbc)
                    cont_end = para_cont.add_run(' and attach that sheet to this form]')
                    cont_end.font.name = 'Arial'
                    cont_end.font.size = Pt(12)

            # Section (b) - format with gold brackets and cream highlight
            # (b) [You have failed to comply...
            # (i) consideration of extension...
            # (ii) enabling a Part 4A certificate to be given.]
            # Section (b) paragraphs based on actual template structure:
            # Para 19: "You have failed to comply..."
            # Para 21: "consideration of extension..."
            # Para 22: "enabling a Part 4A certificate..."
            for i in range(15, min(30, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                txt_lower = txt.lower()

                # "You have failed to comply" - gold bracket at start, cream highlight (section b)
                if "failed to comply" in txt_lower:
                    print(f"DEBUG: Found 'failed to comply' at para {i}: {txt[:50]}")
                    section_b_paras.append(i)
                    para_b = paragraphs[i]
                    para_b_text = txt.strip()
                    for run in para_b.runs:
                        run.text = ""
                    while len(para_b.runs) > 1:
                        para_b._element.remove(para_b.runs[-1]._element)
                    if para_b.runs:
                        para_b.runs[0].text = '['
                        para_b.runs[0].font.bold = True
                        para_b.runs[0].font.color.rgb = BRACKET_COLOR
                    b_content = para_b.add_run(para_b_text)
                    b_content.font.name = 'Arial'
                    b_content.font.size = Pt(12)
                    rPr_bc = b_content._element.get_or_add_rPr()
                    shd_bc = OxmlElement('w:shd')
                    shd_bc.set(qn('w:val'), 'clear')
                    shd_bc.set(qn('w:color'), 'auto')
                    shd_bc.set(qn('w:fill'), 'FFFED5')
                    rPr_bc.append(shd_bc)

                # "consideration of extension" - cream highlight (section b option i)
                elif "consideration of extension" in txt_lower:
                    print(f"DEBUG: Found 'consideration of extension' at para {i}: {txt[:50]}")
                    section_b_i_para = i
                    highlight_yellow(paragraphs[i])

                # "enabling a Part 4A certificate" - cream highlight with gold bracket at end (section b option ii)
                elif "part 4a" in txt_lower:
                    print(f"DEBUG: Found 'Part 4A' at para {i}: {txt[:50]}")
                    section_b_ii_para = i
                    para_b_ii = paragraphs[i]
                    para_b_ii_text = txt.strip()
                    for run in para_b_ii.runs:
                        run.text = ""
                    while len(para_b_ii.runs) > 1:
                        para_b_ii._element.remove(para_b_ii.runs[-1]._element)
                    if para_b_ii.runs:
                        para_b_ii.runs[0].text = para_b_ii_text
                        para_b_ii.runs[0].font.name = 'Arial'
                        para_b_ii.runs[0].font.size = Pt(12)
                        rPr_bii = para_b_ii.runs[0]._element.get_or_add_rPr()
                        shd_bii = OxmlElement('w:shd')
                        shd_bii.set(qn('w:val'), 'clear')
                        shd_bii.set(qn('w:color'), 'auto')
                        shd_bii.set(qn('w:fill'), 'FFFED5')
                        rPr_bii.append(shd_bii)
                    bii_close = para_b_ii.add_run(']')
                    bii_close.font.bold = True
                    bii_close.font.color.rgb = BRACKET_COLOR

            # Debug: Show what was captured
            print(f"DEBUG CTO3: section_a_paras = {section_a_paras}")
            print(f"DEBUG CTO3: section_b_paras = {section_b_paras}")
            print(f"DEBUG CTO3: section_b_i_para = {section_b_i_para}")
            print(f"DEBUG CTO3: section_b_ii_para = {section_b_ii_para}")
            print(f"DEBUG CTO3: reason_a checked = {self.reason_a.isChecked()}")
            print(f"DEBUG CTO3: reason_b checked = {self.reason_b.isChecked()}")
            print(f"DEBUG CTO3: condition_i checked = {self.condition_i_rb.isChecked()}")
            print(f"DEBUG CTO3: condition_ii checked = {self.condition_ii_rb.isChecked()}")

            # Apply strikethrough to the non-selected section
            # If reason_a is selected: strikethrough all of section (b) including (i) and (ii)
            # If reason_b is selected: strikethrough section (a), AND the non-selected (i) or (ii)
            if self.reason_a.isChecked():
                # Strike out section (b) main paragraph
                for idx in section_b_paras:
                    strikethrough_para(paragraphs[idx])
                # Strike out both (i) and (ii)
                if section_b_i_para is not None:
                    strikethrough_para(paragraphs[section_b_i_para])
                if section_b_ii_para is not None:
                    strikethrough_para(paragraphs[section_b_ii_para])
            elif self.reason_b.isChecked():
                # Strike out section (a)
                for idx in section_a_paras:
                    strikethrough_para(paragraphs[idx])
                # Strike out the non-selected (i) or (ii)
                if self.condition_i_rb.isChecked() and section_b_ii_para is not None:
                    # (i) selected, strike out (ii)
                    strikethrough_para(paragraphs[section_b_ii_para])
                elif self.condition_ii_rb.isChecked() and section_b_i_para is not None:
                    # (ii) selected, strike out (i)
                    strikethrough_para(paragraphs[section_b_i_para])

            # Signature line: Signed[ ] Responsible clinician
            for i in range(20, min(35, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if "signed" in txt:
                    para_sig = paragraphs[i]
                    for run in para_sig.runs:
                        run.text = ""
                    while len(para_sig.runs) > 1:
                        para_sig._element.remove(para_sig.runs[-1]._element)
                    if para_sig.runs:
                        para_sig.runs[0].text = 'Signed'
                        para_sig.runs[0].font.name = 'Arial'
                        para_sig.runs[0].font.size = Pt(12)
                    sig_open = para_sig.add_run('[')
                    sig_open.font.bold = True
                    sig_open.font.color.rgb = BRACKET_COLOR
                    rPr_so = sig_open._element.get_or_add_rPr()
                    shd_so = OxmlElement('w:shd')
                    shd_so.set(qn('w:val'), 'clear')
                    shd_so.set(qn('w:color'), 'auto')
                    shd_so.set(qn('w:fill'), 'FFFED5')
                    rPr_so.append(shd_so)
                    sig_content = para_sig.add_run('                                        ')
                    rPr_sc = sig_content._element.get_or_add_rPr()
                    shd_sc = OxmlElement('w:shd')
                    shd_sc.set(qn('w:val'), 'clear')
                    shd_sc.set(qn('w:color'), 'auto')
                    shd_sc.set(qn('w:fill'), 'FFFED5')
                    rPr_sc.append(shd_sc)
                    sig_close = para_sig.add_run(']')
                    sig_close.font.bold = True
                    sig_close.font.color.rgb = BRACKET_COLOR
                    sig_label = para_sig.add_run(' Responsible clinician')
                    sig_label.font.name = 'Arial'
                    sig_label.font.size = Pt(12)
                    break

            # PRINT NAME line
            rc_name = self.rc_name.text().strip()
            for i in range(20, min(35, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "PRINT NAME" in txt:
                    para_pn = paragraphs[i]
                    for run in para_pn.runs:
                        run.text = ""
                    while len(para_pn.runs) > 1:
                        para_pn._element.remove(para_pn.runs[-1]._element)
                    if para_pn.runs:
                        para_pn.runs[0].text = 'PRINT NAME'
                        para_pn.runs[0].font.name = 'Arial'
                        para_pn.runs[0].font.size = Pt(12)
                    pn_open = para_pn.add_run('[')
                    pn_open.font.bold = True
                    pn_open.font.color.rgb = BRACKET_COLOR
                    rPr_pno = pn_open._element.get_or_add_rPr()
                    shd_pno = OxmlElement('w:shd')
                    shd_pno.set(qn('w:val'), 'clear')
                    shd_pno.set(qn('w:color'), 'auto')
                    shd_pno.set(qn('w:fill'), 'FFFED5')
                    rPr_pno.append(shd_pno)
                    pn_content = para_pn.add_run(rc_name if rc_name else '                                        ')
                    rPr_pnc = pn_content._element.get_or_add_rPr()
                    shd_pnc = OxmlElement('w:shd')
                    shd_pnc.set(qn('w:val'), 'clear')
                    shd_pnc.set(qn('w:color'), 'auto')
                    shd_pnc.set(qn('w:fill'), 'FFFED5')
                    rPr_pnc.append(shd_pnc)
                    pn_close = para_pn.add_run(']')
                    pn_close.font.bold = True
                    pn_close.font.color.rgb = BRACKET_COLOR
                    break

            # Date line
            sig_date_str = self.sig_date.date().toString("dd MMMM yyyy")
            for i in range(20, min(35, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if txt.strip().startswith("date"):
                    para_date = paragraphs[i]
                    for run in para_date.runs:
                        run.text = ""
                    while len(para_date.runs) > 1:
                        para_date._element.remove(para_date.runs[-1]._element)
                    if para_date.runs:
                        para_date.runs[0].text = 'Date'
                        para_date.runs[0].font.name = 'Arial'
                        para_date.runs[0].font.size = Pt(12)
                    date_open = para_date.add_run('[')
                    date_open.font.bold = True
                    date_open.font.color.rgb = BRACKET_COLOR
                    rPr_do = date_open._element.get_or_add_rPr()
                    shd_do = OxmlElement('w:shd')
                    shd_do.set(qn('w:val'), 'clear')
                    shd_do.set(qn('w:color'), 'auto')
                    shd_do.set(qn('w:fill'), 'FFFED5')
                    rPr_do.append(shd_do)
                    date_content = para_date.add_run(sig_date_str)
                    rPr_dc = date_content._element.get_or_add_rPr()
                    shd_dc = OxmlElement('w:shd')
                    shd_dc.set(qn('w:val'), 'clear')
                    shd_dc.set(qn('w:color'), 'auto')
                    shd_dc.set(qn('w:fill'), 'FFFED5')
                    rPr_dc.append(shd_dc)
                    date_close = para_date.add_run(']')
                    date_close.font.bold = True
                    date_close.font.color.rgb = BRACKET_COLOR
                    break

            # Time line
            sig_time_str = self.sig_time.time().toString("HH:mm")
            for i in range(20, min(35, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if txt.strip().startswith("time"):
                    para_time = paragraphs[i]
                    for run in para_time.runs:
                        run.text = ""
                    while len(para_time.runs) > 1:
                        para_time._element.remove(para_time.runs[-1]._element)
                    if para_time.runs:
                        para_time.runs[0].text = 'Time'
                        para_time.runs[0].font.name = 'Arial'
                        para_time.runs[0].font.size = Pt(12)
                    time_open = para_time.add_run('[')
                    time_open.font.bold = True
                    time_open.font.color.rgb = BRACKET_COLOR
                    rPr_to = time_open._element.get_or_add_rPr()
                    shd_to = OxmlElement('w:shd')
                    shd_to.set(qn('w:val'), 'clear')
                    shd_to.set(qn('w:color'), 'auto')
                    shd_to.set(qn('w:fill'), 'FFFED5')
                    rPr_to.append(shd_to)
                    time_content = para_time.add_run(sig_time_str)
                    rPr_tc = time_content._element.get_or_add_rPr()
                    shd_tc = OxmlElement('w:shd')
                    shd_tc.set(qn('w:val'), 'clear')
                    shd_tc.set(qn('w:color'), 'auto')
                    shd_tc.set(qn('w:fill'), 'FFFED5')
                    rPr_tc.append(shd_tc)
                    time_close = para_time.add_run(']')
                    time_close.font.bold = True
                    time_close.font.color.rgb = BRACKET_COLOR
                    break

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO3 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        return {
            "age": self.age_spin.value(),
            "gender": "male" if self.gender_male.isChecked() else "female" if self.gender_female.isChecked() else "other" if self.gender_other.isChecked() else "",
            "ethnicity": self.ethnicity_combo.currentText(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "rc_name": self.rc_name.text(),
            "dx_primary": self.dx_primary.currentText(),
            "dx_secondary": self.dx_secondary.currentText(),
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "reason_a": self.reason_a.isChecked(),
            "reason_b": self.reason_b.isChecked(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_neglect": self.self_neglect_cb.isChecked(),
            "self_risky": self.self_risky_cb.isChecked(),
            "self_harm_detail": self.self_harm_detail_cb.isChecked(),
            "others": self.others_cb.isChecked(),
            "others_violence": self.others_violence_cb.isChecked(),
            "others_verbal": self.others_verbal_cb.isChecked(),
            "others_sexual": self.others_sexual_cb.isChecked(),
            "condition_i": self.condition_i_rb.isChecked(),
            "condition_ii": self.condition_ii_rb.isChecked(),
            "grounds": self.grounds_preview.text(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
            "sig_time": self.sig_time.time().toString("HH:mm"),
            "cards": {key: card.get_content() for key, card in self.cards.items()},
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.age_spin.setValue(state.get("age", 0))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        ethnicity = state.get("ethnicity", "Ethnicity")
        idx = self.ethnicity_combo.findText(ethnicity)
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        self.rc_name.setText(state.get("rc_name", ""))
        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)
        dx_secondary = state.get("dx_secondary", "")
        idx = self.dx_secondary.findText(dx_secondary)
        if idx >= 0:
            self.dx_secondary.setCurrentIndex(idx)
        else:
            self.dx_secondary.setCurrentText(dx_secondary)
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.reason_a.setChecked(state.get("reason_a", False))
        self.reason_b.setChecked(state.get("reason_b", False))
        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_neglect_cb.setChecked(state.get("self_neglect", False))
        self.self_risky_cb.setChecked(state.get("self_risky", False))
        self.self_harm_detail_cb.setChecked(state.get("self_harm_detail", False))
        self.others_cb.setChecked(state.get("others", False))
        self.others_violence_cb.setChecked(state.get("others_violence", False))
        self.others_verbal_cb.setChecked(state.get("others_verbal", False))
        self.others_sexual_cb.setChecked(state.get("others_sexual", False))
        self.condition_i_rb.setChecked(state.get("condition_i", False))
        self.condition_ii_rb.setChecked(state.get("condition_ii", False))
        self.grounds_preview.setText(state.get("grounds", ""))
        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))
        if state.get("sig_time"):
            self.sig_time.setTime(QTime.fromString(state["sig_time"], "HH:mm"))
        # Restore card contents
        cards_state = state.get("cards", {})
        for key, content in cards_state.items():
            if key in self.cards:
                self.cards[key].set_content_text(content)

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[CTO3Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[CTO3Form] Set gender: {gender}")
        # Set age if not already set (from age field or calculate from DOB)
        if hasattr(self, 'age_spin') and self.age_spin.value() == 0:
            age = patient_info.get("age")
            if not age and patient_info.get("dob"):
                from datetime import datetime
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if age and 0 < age < 120:
                self.age_spin.setValue(age)
                print(f"[CTO3Form] Set age: {age}")
        # Set ethnicity if not already selected
        if patient_info.get("ethnicity") and hasattr(self, 'ethnicity_combo'):
            current = self.ethnicity_combo.currentText()
            if current in ("Ethnicity", "Not specified", "Select..."):
                idx = self.ethnicity_combo.findText(patient_info["ethnicity"], Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    self.ethnicity_combo.setCurrentIndex(idx)
                    print(f"[CTO3Form] Set ethnicity: {patient_info['ethnicity']}")
