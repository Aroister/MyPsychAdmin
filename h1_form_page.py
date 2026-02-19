# ================================================================
#  H1 FORM PAGE — Section 5(2) Report on Hospital In-patient
#  Mental Health Act 1983 - Form H1 Regulation 4(1)(g)
#  Part 1 only - Holding power report
#  CARD/POPUP LAYOUT with ResizableSection
# ================================================================

from __future__ import annotations

from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QGroupBox, QToolButton, QRadioButton,
    QButtonGroup, QComboBox, QCompleter, QStackedWidget, QSplitter,
    QStyleFactory
)

from background_history_popup import ResizableSection
from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor
from utils.resource_path import resource_path

# ICD-10 data - curated list matching iOS app
try:
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []


# ================================================================
# H1 CARD WIDGET
# ================================================================
class H1CardWidget(QFrame):
    """Card widget with fixed header and scrollable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("h1Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#h1Card {
                background: rgba(255,255,255,0.65);
                border: 1px solid rgba(0,0,0,0.08);
                border-radius: 18px;
            }
            QFrame#h1Card:hover {
                border-color: #1e3a5f;
                background: #eff6ff;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QLabel#cardTitle {
                font-size: 20px;
                font-weight: 600;
                color: #1e3a5f;
                padding-bottom: 4px;
            }
            QFrame#divider {
                background: rgba(0,0,0,0.10);
                height: 1px;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 12)
        layout.setSpacing(0)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setObjectName("cardTitle")
        self.title_label.setFixedHeight(24)
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        # Divider
        divider = QFrame()
        divider.setObjectName("divider")
        divider.setFixedHeight(1)
        layout.addWidget(divider)

        # Scrollable content
        self.content_scroll = QScrollArea()
        self.content_scroll.setWidgetResizable(True)
        self.content_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        self.content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.content_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        content_widget = QWidget()
        content_widget.setStyleSheet("background: transparent;")
        self.content_layout = QVBoxLayout(content_widget)
        self.content_layout.setContentsMargins(0, 8, 0, 4)
        self.content_layout.setSpacing(0)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 16px;
                color: #374151;
                padding: 4px;
                background: transparent;
                border: none;
            }
        """)
        self.content.setFrameShape(QFrame.Shape.NoFrame)
        self.content_layout.addWidget(self.content)
        self.content_layout.addStretch()

        self.content_scroll.setWidget(content_widget)
        layout.addWidget(self.content_scroll, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def set_content(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)


# ================================================================
# TOOLBAR
# ================================================================
class H1Toolbar(QWidget):
    """Toolbar for the H1 Form Page."""

    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            H1Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN H1 FORM PAGE
# ================================================================
class H1FormPage(QWidget):
    """Page for completing MHA Form H1 - Section 5(2) Report on Hospital In-patient."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}
        self._reasons_narrative = ""

        self._setup_ui()
        self._prefill_practitioner()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill_practitioner(self):
        """Pre-fill practitioner details from saved settings."""
        if self._my_details:
            self.prac_name.setText(self._my_details.get("full_name", "") or "")
            self.hospital_name.setText(self._my_details.get("hospital_name", "") or "")
            self.hospital_address.setText(self._my_details.get("hospital_address", "") or "")
            # Update cards after prefill
            self._update_practitioner_preview()
            self._update_hospital_preview()

    def _setup_ui(self):
        self.setStyleSheet("background: #f3f4f6;")

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #1e3a5f; border-bottom: 1px solid #163049;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QToolButton()
        back_btn.setText("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QToolButton {
                background: rgba(255,255,255,0.15);
                color: white;
                font-size: 17px;
                font-weight: 600;
                border: none;
                border-radius: 6px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form H1 — Section 5(2) Report on Hospital In-patient")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - splitter with cards on left, popup on right
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(6)
        splitter.setStyleSheet("""
            QSplitter { background: #f9fafb; }
            QSplitter::handle { background: #d1d5db; }
            QSplitter::handle:hover { background: #6BAF8D; }
        """)

        # Left side - Cards panel
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: #f9fafb; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: #f9fafb;")
        cards_layout = QVBoxLayout(cards_container)
        cards_layout.setContentsMargins(24, 24, 24, 24)
        cards_layout.setSpacing(16)

        # Create cards with ResizableSection
        self.cards["patient"] = self._create_card("Patient", "patient", 110)
        self.cards["hospital"] = self._create_card("Hospital", "hospital", 110)
        self.cards["practitioner"] = self._create_card("Practitioner", "practitioner", 130)
        self.cards["reasons"] = self._create_card("Reasons for Detention", "reasons", 150)
        self.cards["delivery"] = self._create_card("Delivery & Signature", "delivery", 110)

        for key in ["patient", "hospital", "practitioner", "reasons", "delivery"]:
            cards_layout.addWidget(self.cards[key]["section"])

        cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        splitter.addWidget(cards_scroll)

        # Right side - Popup panel
        self.popup_panel = QFrame()
        self.popup_panel.setStyleSheet("QFrame { background: rgba(245,245,245,0.95); border-left: 1px solid rgba(0,0,0,0.08); }")
        panel_layout = QVBoxLayout(self.popup_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setStyleSheet("""
            font-size: 22px; font-weight: 700; color: #1e3a5f;
            background: rgba(219,234,254,0.85); padding: 8px 12px; border-radius: 8px;
        """)
        panel_layout.addWidget(self.panel_title)

        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        # Create popups
        self.popup_stack.addWidget(self._create_patient_popup())
        self.popup_stack.addWidget(self._create_hospital_popup())
        self.popup_stack.addWidget(self._create_practitioner_popup())
        self.popup_stack.addWidget(self._create_reasons_popup())
        self.popup_stack.addWidget(self._create_delivery_popup())

        # Initialize cards with default date values
        self._update_delivery_preview()

        splitter.addWidget(self.popup_panel)

        # Set initial sizes (cards: 350, popup: rest)
        splitter.setSizes([350, 500])

        main_layout.addWidget(splitter, 1)

    def _create_card(self, title: str, key: str, height: int):
        """Create a card with ResizableSection wrapper."""
        section = ResizableSection()
        section.set_content_height(height + 60)
        section._min_height = 120
        section._max_height = 400

        card = H1CardWidget(title, key)
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        section.set_content(card)

        return {"section": section, "card": card}

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"patient": 0, "hospital": 1, "practitioner": 2, "reasons": 3, "delivery": 4}
        title_map = {"patient": "Patient Details", "hospital": "Hospital Details", "practitioner": "Practitioner Details", "reasons": "Reasons for Detention", "delivery": "Delivery & Signature"}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            self.panel_title.setText(title_map.get(key, ""))
            # Highlight selected card
            for k, data in self.cards.items():
                card = data["card"]
                if k == key:
                    card.setStyleSheet("""
                        QFrame#h1Card {
                            background: rgba(239,246,255,0.85);
                            border: 2px solid #1e3a5f;
                            border-radius: 18px;
                        }
                        QLabel#cardTitle {
                            font-size: 20px;
                            font-weight: 600;
                            color: #1e3a5f;
                            padding-bottom: 4px;
                        }
                        QFrame#divider {
                            background: rgba(0,0,0,0.10);
                            height: 1px;
                        }
                    """)
                else:
                    card.setStyleSheet("""
                        QFrame#h1Card {
                            background: rgba(255,255,255,0.65);
                            border: 1px solid rgba(0,0,0,0.08);
                            border-radius: 18px;
                        }
                        QFrame#h1Card:hover {
                            border-color: #1e3a5f;
                            background: #eff6ff;
                        }
                        QLabel#cardTitle {
                            font-size: 20px;
                            font-weight: 600;
                            color: #1e3a5f;
                            padding-bottom: 4px;
                        }
                        QFrame#divider {
                            background: rgba(0,0,0,0.10);
                            height: 1px;
                        }
                    """)

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_patient_popup(self):
        """Create the patient details popup."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Patient name
        name_lbl = QLabel("Full Name")
        name_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f;")
        layout.addWidget(name_lbl)

        self.patient_name = QLineEdit()
        self.patient_name.setPlaceholderText("Full name of patient")
        self.patient_name.setStyleSheet("padding: 10px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 17px;")
        layout.addWidget(self.patient_name)

        # Gender selection
        gender_lbl = QLabel("Gender")
        gender_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f;")
        layout.addWidget(gender_lbl)

        gender_row = QHBoxLayout()
        gender_row.setSpacing(16)

        self.gender_group = QButtonGroup(self)

        self.gender_male = QRadioButton("Male")
        self.gender_male.setStyleSheet("font-size: 17px; color: #374151;")
        self.gender_male.toggled.connect(self._build_narrative)
        self.gender_group.addButton(self.gender_male)
        gender_row.addWidget(self.gender_male)

        self.gender_female = QRadioButton("Female")
        self.gender_female.setStyleSheet("font-size: 17px; color: #374151;")
        self.gender_female.toggled.connect(self._build_narrative)
        self.gender_group.addButton(self.gender_female)
        gender_row.addWidget(self.gender_female)

        self.gender_other = QRadioButton("Other")
        self.gender_other.setStyleSheet("font-size: 17px; color: #374151;")
        self.gender_other.toggled.connect(self._build_narrative)
        self.gender_group.addButton(self.gender_other)
        gender_row.addWidget(self.gender_other)

        gender_row.addStretch()
        layout.addLayout(gender_row)

        layout.addStretch()

        # Auto-sync to card
        self.patient_name.textChanged.connect(self._update_patient_preview)
        self.gender_male.toggled.connect(self._update_patient_preview)
        self.gender_female.toggled.connect(self._update_patient_preview)
        self.gender_other.toggled.connect(self._update_patient_preview)

        popup.setWidget(container)
        return popup

    def _create_hospital_popup(self):
        """Create the hospital details popup."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Hospital name
        name_lbl = QLabel("Hospital Name")
        name_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f;")
        layout.addWidget(name_lbl)

        self.hospital_name = QLineEdit()
        self.hospital_name.setPlaceholderText("Hospital name")
        self.hospital_name.setStyleSheet("padding: 10px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 17px;")
        layout.addWidget(self.hospital_name)

        # Hospital address
        addr_lbl = QLabel("Hospital Address")
        addr_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f;")
        layout.addWidget(addr_lbl)

        self.hospital_address = QLineEdit()
        self.hospital_address.setPlaceholderText("Hospital address")
        self.hospital_address.setStyleSheet("padding: 10px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 17px;")
        layout.addWidget(self.hospital_address)

        layout.addStretch()

        # Auto-sync to card
        self.hospital_name.textChanged.connect(self._update_hospital_preview)
        self.hospital_address.textChanged.connect(self._update_hospital_preview)

        popup.setWidget(container)
        return popup

    def _create_practitioner_popup(self):
        """Create the practitioner details popup."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Practitioner name
        name_lbl = QLabel("Full Name")
        name_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f;")
        layout.addWidget(name_lbl)

        self.prac_name = QLineEdit()
        self.prac_name.setPlaceholderText("Full name")
        self.prac_name.setStyleSheet("padding: 10px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 17px;")
        layout.addWidget(self.prac_name)

        # Clinician type (first choice)
        clin_lbl = QLabel("Clinician type:")
        clin_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f; margin-top: 8px;")
        layout.addWidget(clin_lbl)

        self.clinician_type_group = QButtonGroup(self)

        self.clinician_rmp = QRadioButton("Registered medical practitioner")
        self.clinician_rmp.setStyleSheet("font-size: 17px; color: #374151;")
        self.clinician_rmp.setChecked(True)
        self.clinician_type_group.addButton(self.clinician_rmp)
        layout.addWidget(self.clinician_rmp)

        self.clinician_ac = QRadioButton("Approved clinician (not a registered medical practitioner)")
        self.clinician_ac.setStyleSheet("font-size: 17px; color: #374151;")
        self.clinician_type_group.addButton(self.clinician_ac)
        layout.addWidget(self.clinician_ac)

        # Nominee (second choice)
        nominee_lbl = QLabel("Are you a nominee?")
        nominee_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e3a5f; margin-top: 8px;")
        layout.addWidget(nominee_lbl)

        self.nominee_group = QButtonGroup(self)

        self.nominee_no = QRadioButton("No - I am in charge of the treatment of this patient")
        self.nominee_no.setStyleSheet("font-size: 17px; color: #374151;")
        self.nominee_no.setChecked(True)
        self.nominee_group.addButton(self.nominee_no)
        layout.addWidget(self.nominee_no)

        self.nominee_yes = QRadioButton("Yes - I am a nominee of the person in charge of treatment")
        self.nominee_yes.setStyleSheet("font-size: 17px; color: #374151;")
        self.nominee_group.addButton(self.nominee_yes)
        layout.addWidget(self.nominee_yes)

        # Keep old attributes for compatibility
        self.prac_in_charge = self.nominee_no
        self.prac_nominee = self.nominee_yes

        layout.addStretch()

        # Auto-sync to card
        self.prac_name.textChanged.connect(self._update_practitioner_preview)
        self.nominee_no.toggled.connect(self._update_practitioner_preview)
        self.nominee_yes.toggled.connect(self._update_practitioner_preview)
        self.clinician_rmp.toggled.connect(self._update_practitioner_preview)
        self.clinician_ac.toggled.connect(self._update_practitioner_preview)

        popup.setWidget(container)
        return popup

    def _create_reasons_popup(self):
        """Create the reasons for detention popup with live preview."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        container = QWidget()
        container.setMaximumWidth(500)
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Hidden label for state storage
        self.reasons_preview = QLabel("")
        self.reasons_preview.hide()

        # Controls frame
        controls_frame = QFrame()
        controls_frame.setStyleSheet("QFrame { background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 8px; }")
        controls_layout = QVBoxLayout(controls_frame)
        controls_layout.setContentsMargins(16, 12, 16, 12)
        controls_layout.setSpacing(12)

        # Diagnosis dropdown
        dx_lbl = QLabel("Mental Disorder (ICD-10)")
        dx_lbl.setStyleSheet("font-size: 16px; font-weight: 700; color: #166534;")
        controls_layout.addWidget(dx_lbl)

        self.dx_primary = QComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 16px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_primary.currentTextChanged.connect(self._build_narrative)
        controls_layout.addWidget(self.dx_primary)

        # Reasons checkboxes
        reasons_lbl = QLabel("Reasons")
        reasons_lbl.setStyleSheet("font-size: 16px; font-weight: 700; color: #1e40af;")
        controls_layout.addWidget(reasons_lbl)

        self.cb_refusing = QCheckBox("Refusing to remain in hospital")
        self.cb_refusing.setStyleSheet("font-size: 16px; color: #374151;")
        self.cb_refusing.stateChanged.connect(self._build_narrative)
        controls_layout.addWidget(self.cb_refusing)

        self.cb_very_unwell = QCheckBox("Very unwell")
        self.cb_very_unwell.setStyleSheet("font-size: 16px; color: #374151;")
        self.cb_very_unwell.stateChanged.connect(self._build_narrative)
        controls_layout.addWidget(self.cb_very_unwell)

        self.cb_acute = QCheckBox("Acute deterioration")
        self.cb_acute.setStyleSheet("font-size: 16px; color: #374151;")
        self.cb_acute.stateChanged.connect(self._build_narrative)
        controls_layout.addWidget(self.cb_acute)

        # Risk section
        risk_lbl = QLabel("Significant risk to:")
        risk_lbl.setStyleSheet("font-size: 16px; font-weight: 700; color: #991b1b;")
        controls_layout.addWidget(risk_lbl)

        risk_row = QHBoxLayout()
        risk_row.setSpacing(16)

        self.cb_risk_self = QCheckBox("Self")
        self.cb_risk_self.setStyleSheet("font-size: 16px; color: #374151;")
        self.cb_risk_self.stateChanged.connect(self._build_narrative)
        risk_row.addWidget(self.cb_risk_self)

        self.cb_risk_others = QCheckBox("Others")
        self.cb_risk_others.setStyleSheet("font-size: 16px; color: #374151;")
        self.cb_risk_others.stateChanged.connect(self._build_narrative)
        risk_row.addWidget(self.cb_risk_others)

        risk_row.addStretch()
        controls_layout.addLayout(risk_row)

        layout.addWidget(controls_frame)

        layout.addStretch()
        popup.setWidget(container)
        return popup

    def _create_delivery_popup(self):
        """Create the delivery method and signature popup."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Delivery method section
        delivery_lbl = QLabel("Report Furnishing Method")
        delivery_lbl.setStyleSheet("font-size: 17px; font-weight: 700; color: #854d0e;")
        layout.addWidget(delivery_lbl)

        self.delivery_group = QButtonGroup(self)

        self.delivery_internal = QRadioButton("(a) Internal mail system")
        self.delivery_internal.setStyleSheet("font-size: 17px; color: #374151;")
        self.delivery_internal.setChecked(True)
        self.delivery_internal.toggled.connect(self._on_delivery_changed)
        self.delivery_group.addButton(self.delivery_internal)
        layout.addWidget(self.delivery_internal)

        # Time for internal mail
        time_row = QHBoxLayout()
        time_row.setContentsMargins(20, 0, 0, 0)
        time_lbl = QLabel("Time consigned:")
        time_lbl.setStyleSheet("font-size: 16px; color: #6b7280;")
        time_row.addWidget(time_lbl)
        self.internal_time = QTimeEdit()
        self.internal_time.setTime(QTime.currentTime())
        self.internal_time.setDisplayFormat("HH:mm")
        self.internal_time.setStyleSheet("padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;")
        time_row.addWidget(self.internal_time)
        time_row.addStretch()
        layout.addLayout(time_row)

        self.delivery_electronic = QRadioButton("(b) Electronic communication")
        self.delivery_electronic.setStyleSheet("font-size: 17px; color: #374151;")
        self.delivery_group.addButton(self.delivery_electronic)
        layout.addWidget(self.delivery_electronic)

        self.delivery_hand = QRadioButton("(c) Delivered by hand")
        self.delivery_hand.setStyleSheet("font-size: 17px; color: #374151;")
        self.delivery_group.addButton(self.delivery_hand)
        layout.addWidget(self.delivery_hand)

        # Signature section
        sig_lbl = QLabel("Signature")
        sig_lbl.setStyleSheet("font-size: 17px; font-weight: 700; color: #1e3a5f; margin-top: 16px;")
        layout.addWidget(sig_lbl)

        date_row = QHBoxLayout()
        date_label = QLabel("Date:")
        date_label.setStyleSheet("font-size: 16px; font-weight: 500; color: #374151;")
        date_row.addWidget(date_label)

        self.sig_date = QDateEdit()
        self.sig_date.setDate(QDate.currentDate())
        self.sig_date.setCalendarPopup(True)
        self.sig_date.setDisplayFormat("dd/MM/yyyy")
        self.sig_date.setStyleSheet("padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;")
        date_row.addWidget(self.sig_date)
        date_row.addStretch()
        layout.addLayout(date_row)

        layout.addStretch()

        # Auto-sync to card
        self.delivery_internal.toggled.connect(self._update_delivery_preview)
        self.delivery_electronic.toggled.connect(self._update_delivery_preview)
        self.delivery_hand.toggled.connect(self._update_delivery_preview)
        self.internal_time.timeChanged.connect(self._update_delivery_preview)
        self.sig_date.dateChanged.connect(self._update_delivery_preview)

        popup.setWidget(container)
        return popup

    # ----------------------------------------------------------------
    # PREVIEW UPDATES
    # ----------------------------------------------------------------
    def _update_all_previews(self):
        self._update_patient_preview()
        self._update_hospital_preview()
        self._update_practitioner_preview()
        self._update_reasons_preview()
        self._update_delivery_preview()

    def _update_patient_preview(self):
        lines = []
        if self.patient_name.text():
            lines.append(self.patient_name.text())
        self.cards["patient"]["card"].set_content("\n".join(lines))

    def _update_hospital_preview(self):
        lines = []
        if self.hospital_name.text():
            lines.append(self.hospital_name.text())
        if self.hospital_address.text():
            lines.append(self.hospital_address.text())
        self.cards["hospital"]["card"].set_content("\n".join(lines))

    def _update_practitioner_preview(self):
        lines = []
        if self.prac_name.text():
            lines.append(self.prac_name.text())
        if self.prac_in_charge.isChecked():
            lines.append("In charge of treatment")
        else:
            lines.append("Nominee")
        if self.clinician_rmp.isChecked():
            lines.append("Registered medical practitioner")
        else:
            lines.append("Approved clinician")
        self.cards["practitioner"]["card"].set_content("\n".join(lines))

    def _update_reasons_preview(self):
        text = self._reasons_narrative
        self.cards["reasons"]["card"].set_content(text if text else "")

    def _update_delivery_preview(self):
        lines = []
        if self.delivery_internal.isChecked():
            lines.append(self.internal_time.time().toString('HH:mm'))
        lines.append(self.sig_date.date().toString('dd MMM yyyy'))
        self.cards["delivery"]["card"].set_content("\n".join(lines))

    def _on_delivery_changed(self):
        """Enable/disable time field based on delivery method."""
        self.internal_time.setEnabled(self.delivery_internal.isChecked())

    def _build_narrative(self):
        """Build narrative text from diagnosis and tick boxes."""
        parts = []

        # Determine pronouns based on gender
        if self.gender_male.isChecked():
            subj, poss, reflex = "He", "His", "himself"
        elif self.gender_female.isChecked():
            subj, poss, reflex = "She", "Her", "herself"
        elif self.gender_other.isChecked():
            subj, poss, reflex = "They", "Their", "themselves"
        else:
            subj, poss, reflex = "They", "Their", "themselves"

        is_verb = "is" if subj != "They" else "are"

        dx_text = self.dx_primary.currentText().strip()
        if dx_text:
            parts.append(f"The patient suffers from {dx_text}.")

        if self.cb_refusing.isChecked():
            parts.append(f"{subj} {is_verb} refusing to remain in hospital informally.")

        if self.cb_very_unwell.isChecked() and self.cb_acute.isChecked():
            parts.append(f"{subj} {is_verb} very unwell and suffering an acute deterioration of {poss.lower()} mental state.")
        elif self.cb_very_unwell.isChecked():
            parts.append(f"{subj} {is_verb} currently very unwell.")
        elif self.cb_acute.isChecked():
            parts.append(f"{subj} {is_verb} suffering an acute deterioration of {poss.lower()} mental state.")

        if self.cb_risk_self.isChecked() and self.cb_risk_others.isChecked():
            parts.append(f"{poss} risk to {reflex} and others is significant warranting a mental health act assessment.")
        elif self.cb_risk_self.isChecked():
            parts.append(f"{poss} risk to {reflex} is significant warranting a mental health act assessment.")
        elif self.cb_risk_others.isChecked():
            parts.append(f"{poss} risk to others is significant warranting a mental health act assessment.")

        self._reasons_narrative = " ".join(parts)
        self.reasons_preview.setText(self._reasons_narrative)
        # Auto-sync to card
        self._update_reasons_preview()

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        """Clear all form fields."""
        reply = QMessageBox.question(
            self, "Clear Form",
            "Are you sure you want to clear all fields?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.patient_name.clear()
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.dx_primary.setCurrentIndex(0)
            self.cb_refusing.setChecked(False)
            self.cb_very_unwell.setChecked(False)
            self.cb_risk_self.setChecked(False)
            self.cb_risk_others.setChecked(False)
            self.cb_acute.setChecked(False)
            self._reasons_narrative = ""
            self.reasons_preview.clear()
            self.prac_in_charge.setChecked(True)
            self.clinician_rmp.setChecked(True)
            self.delivery_internal.setChecked(True)
            self.internal_time.setTime(QTime.currentTime())
            self.sig_date.setDate(QDate.currentDate())
            self._prefill_practitioner()
            self._update_all_previews()

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form H1",
            f"Form_H1_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_H1_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form H1 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color (#918C0D) and cream highlight (#FFFED5)
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)
            CREAM_FILL = 'FFFED5'

            # Pre-process: clean ALL paragraphs - remove permission markers, convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), CREAM_FILL)
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), CREAM_FILL)

            def set_entry_box(para, content: str):
                """Set entry box with gold brackets - content between brackets is cream."""
                if not content or not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Opening bracket
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def set_labeled_entry(para, label: str, content: str, suffix: str = ""):
                """Set labeled entry - only content between brackets is cream."""
                if not content or not content.strip():
                    content = '                                        '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Label
                if para.runs:
                    para.runs[0].text = label
                    label_run = para.runs[0]
                else:
                    label_run = para.add_run(label)
                label_run.font.name = 'Arial'
                label_run.font.size = Pt(12)
                # Opening bracket
                bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)
                # Suffix
                if suffix:
                    suffix_run = para.add_run(suffix)
                    suffix_run.font.name = 'Arial'
                    suffix_run.font.size = Pt(12)

            def format_option_first(para, content: str, strike: bool = False):
                """First option - opening bracket only."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Opening bracket
                if para.runs:
                    para.runs[0].text = '['
                    ob = para.runs[0]
                else:
                    ob = para.add_run('[')
                ob.font.name = 'Arial'
                ob.font.size = Pt(12)
                ob.font.bold = True
                ob.font.color.rgb = BRACKET_COLOR
                rPr_ob = ob._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                if strike:
                    ob.font.strike = True
                # Content
                ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True

            def format_option_middle(para, content: str, strike: bool = False):
                """Middle option - no brackets, just cream highlight."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Content only
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True

            def format_option_last(para, content: str, strike: bool = False):
                """Last option - closing bracket only."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Content
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True
                # Closing bracket
                cb = para.add_run(']')
                cb.font.name = 'Arial'
                cb.font.size = Pt(12)
                cb.font.bold = True
                cb.font.color.rgb = BRACKET_COLOR
                rPr_cb = cb._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)
                if strike:
                    cb.font.strike = True

            paragraphs = doc.paragraphs

            # Hospital - goes in para 5 (use card content)
            hospital_text = self.cards["hospital"]["card"].get_content().replace("\n", ", ")
            set_entry_box(paragraphs[5], hospital_text)

            # Practitioner - goes in para 7 (name only)
            prac_name = self.prac_name.text().strip()
            set_entry_box(paragraphs[7], prac_name)

            # (a) and (b) options - Para 9 = (a), Para 10 = (b)
            # Determine which to strikethrough
            strike_a = not self.nominee_no.isChecked() if hasattr(self, 'nominee_no') else False
            strike_b = self.nominee_no.isChecked() if hasattr(self, 'nominee_no') else True

            # Format (a) option with opening bracket
            format_option_first(paragraphs[9],
                "the registered medical practitioner/the approved clinician (who is not a registered medical practitioner)<delete the phrase which does not apply>",
                strike=strike_a)

            # Format (b) option with closing bracket
            format_option_last(paragraphs[10],
                "a registered medical practitioner/an approved clinician (who is not a registered medical practitioner)* who is the nominee of the registered medical practitioner or approved clinician (who is not a registered medical practitioner) <*delete the phrase which does not apply>",
                strike=strike_b)

            # Patient - goes in para 12 (use card content)
            patient_text = self.cards["patient"]["card"].get_content().replace("\n", ", ")
            set_entry_box(paragraphs[12], patient_text)

            # Reasons - goes in para 15-17 (uses card content)
            # Para 15 gets opening bracket + content, para 16-17 get cream + closing bracket on 17
            reasons_text = self.cards["reasons"]["card"].get_content().strip()
            if not reasons_text:
                reasons_text = '                                                                                                    '

            # Para 15 - opening bracket + content
            for run in paragraphs[15].runs:
                run.text = ""
            while len(paragraphs[15].runs) > 1:
                paragraphs[15]._element.remove(paragraphs[15].runs[-1]._element)
            pPr = paragraphs[15]._element.get_or_add_pPr()
            for old_shd in pPr.findall(qn('w:shd')):
                pPr.remove(old_shd)
            if paragraphs[15].runs:
                paragraphs[15].runs[0].text = '['
                ob = paragraphs[15].runs[0]
            else:
                ob = paragraphs[15].add_run('[')
            ob.font.name = 'Arial'
            ob.font.size = Pt(12)
            ob.font.bold = True
            ob.font.color.rgb = BRACKET_COLOR
            rPr_ob = ob._element.get_or_add_rPr()
            shd_ob = OxmlElement('w:shd')
            shd_ob.set(qn('w:val'), 'clear')
            shd_ob.set(qn('w:color'), 'auto')
            shd_ob.set(qn('w:fill'), CREAM_FILL)
            rPr_ob.append(shd_ob)
            ct = paragraphs[15].add_run(reasons_text)
            ct.font.name = 'Arial'
            ct.font.size = Pt(12)
            rPr_ct = ct._element.get_or_add_rPr()
            shd_ct = OxmlElement('w:shd')
            shd_ct.set(qn('w:val'), 'clear')
            shd_ct.set(qn('w:color'), 'auto')
            shd_ct.set(qn('w:fill'), CREAM_FILL)
            rPr_ct.append(shd_ct)

            # Para 16 - cream continuation
            for run in paragraphs[16].runs:
                run.text = ""
            while len(paragraphs[16].runs) > 1:
                paragraphs[16]._element.remove(paragraphs[16].runs[-1]._element)
            pPr16 = paragraphs[16]._element.get_or_add_pPr()
            for old_shd in pPr16.findall(qn('w:shd')):
                pPr16.remove(old_shd)
            if paragraphs[16].runs:
                paragraphs[16].runs[0].text = '                                                                                                    '
                ct16 = paragraphs[16].runs[0]
            else:
                ct16 = paragraphs[16].add_run('                                                                                                    ')
            ct16.font.name = 'Arial'
            ct16.font.size = Pt(12)
            rPr16 = ct16._element.get_or_add_rPr()
            shd16 = OxmlElement('w:shd')
            shd16.set(qn('w:val'), 'clear')
            shd16.set(qn('w:color'), 'auto')
            shd16.set(qn('w:fill'), CREAM_FILL)
            rPr16.append(shd16)

            # Para 17 - cream + closing bracket
            for run in paragraphs[17].runs:
                run.text = ""
            while len(paragraphs[17].runs) > 1:
                paragraphs[17]._element.remove(paragraphs[17].runs[-1]._element)
            pPr17 = paragraphs[17]._element.get_or_add_pPr()
            for old_shd in pPr17.findall(qn('w:shd')):
                pPr17.remove(old_shd)
            if paragraphs[17].runs:
                paragraphs[17].runs[0].text = '                                                                                                    '
                ct17 = paragraphs[17].runs[0]
            else:
                ct17 = paragraphs[17].add_run('                                                                                                    ')
            ct17.font.name = 'Arial'
            ct17.font.size = Pt(12)
            rPr17 = ct17._element.get_or_add_rPr()
            shd17 = OxmlElement('w:shd')
            shd17.set(qn('w:val'), 'clear')
            shd17.set(qn('w:color'), 'auto')
            shd17.set(qn('w:fill'), CREAM_FILL)
            rPr17.append(shd17)
            cb17 = paragraphs[17].add_run(']')
            cb17.font.name = 'Arial'
            cb17.font.size = Pt(12)
            cb17.font.bold = True
            cb17.font.color.rgb = BRACKET_COLOR
            rPr_cb17 = cb17._element.get_or_add_rPr()
            shd_cb17 = OxmlElement('w:shd')
            shd_cb17.set(qn('w:val'), 'clear')
            shd_cb17.set(qn('w:color'), 'auto')
            shd_cb17.set(qn('w:fill'), CREAM_FILL)
            rPr_cb17.append(shd_cb17)

            # Part 1 Delivery method options - Para 20, 21, 22
            any_delivery_selected = self.delivery_internal.isChecked() or self.delivery_electronic.isChecked() or self.delivery_hand.isChecked()

            # Para 20 - first option with opening bracket
            time_str = self.internal_time.time().toString("HH:mm") if self.delivery_internal.isChecked() else "[time]"
            strike_20 = any_delivery_selected and not self.delivery_internal.isChecked()
            format_option_first(paragraphs[20],
                f"consigning it to the hospital managers' internal mail system today at {time_str}",
                strike=strike_20)

            # Para 21 - middle option
            strike_21 = any_delivery_selected and not self.delivery_electronic.isChecked()
            format_option_middle(paragraphs[21],
                "today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication",
                strike=strike_21)

            # Para 22 - last option with closing bracket
            strike_22 = any_delivery_selected and not self.delivery_hand.isChecked()
            format_option_last(paragraphs[22],
                "delivering it (or having it delivered) by hand to a person authorised by the hospital managers to receive it.",
                strike=strike_22)

            # Part 1 Signature/Date - Para 23
            sig_date = self.sig_date.date().toString("dd/MM/yyyy")
            for run in paragraphs[23].runs:
                run.text = ""
            while len(paragraphs[23].runs) > 1:
                paragraphs[23]._element.remove(paragraphs[23].runs[-1]._element)
            pPr23 = paragraphs[23]._element.get_or_add_pPr()
            for old_shd in pPr23.findall(qn('w:shd')):
                pPr23.remove(old_shd)
            # Signed label
            if paragraphs[23].runs:
                paragraphs[23].runs[0].text = 'Signed'
                sl = paragraphs[23].runs[0]
            else:
                sl = paragraphs[23].add_run('Signed')
            sl.font.name = 'Arial'
            sl.font.size = Pt(12)
            # Signed placeholder
            sob = paragraphs[23].add_run('[')
            sob.font.bold = True
            sob.font.color.rgb = BRACKET_COLOR
            rPr_sob = sob._element.get_or_add_rPr()
            shd_sob = OxmlElement('w:shd')
            shd_sob.set(qn('w:val'), 'clear')
            shd_sob.set(qn('w:color'), 'auto')
            shd_sob.set(qn('w:fill'), CREAM_FILL)
            rPr_sob.append(shd_sob)
            sct = paragraphs[23].add_run('                                        ')
            rPr_sct = sct._element.get_or_add_rPr()
            shd_sct = OxmlElement('w:shd')
            shd_sct.set(qn('w:val'), 'clear')
            shd_sct.set(qn('w:color'), 'auto')
            shd_sct.set(qn('w:fill'), CREAM_FILL)
            rPr_sct.append(shd_sct)
            scb = paragraphs[23].add_run(']')
            scb.font.bold = True
            scb.font.color.rgb = BRACKET_COLOR
            rPr_scb = scb._element.get_or_add_rPr()
            shd_scb = OxmlElement('w:shd')
            shd_scb.set(qn('w:val'), 'clear')
            shd_scb.set(qn('w:color'), 'auto')
            shd_scb.set(qn('w:fill'), CREAM_FILL)
            rPr_scb.append(shd_scb)
            # Date label
            dl = paragraphs[23].add_run(' Date')
            dl.font.name = 'Arial'
            dl.font.size = Pt(12)
            # Date placeholder
            dob = paragraphs[23].add_run('[')
            dob.font.bold = True
            dob.font.color.rgb = BRACKET_COLOR
            rPr_dob = dob._element.get_or_add_rPr()
            shd_dob = OxmlElement('w:shd')
            shd_dob.set(qn('w:val'), 'clear')
            shd_dob.set(qn('w:color'), 'auto')
            shd_dob.set(qn('w:fill'), CREAM_FILL)
            rPr_dob.append(shd_dob)
            dct = paragraphs[23].add_run(f'                    {sig_date}                    ')
            rPr_dct = dct._element.get_or_add_rPr()
            shd_dct = OxmlElement('w:shd')
            shd_dct.set(qn('w:val'), 'clear')
            shd_dct.set(qn('w:color'), 'auto')
            shd_dct.set(qn('w:fill'), CREAM_FILL)
            rPr_dct.append(shd_dct)
            dcb = paragraphs[23].add_run(']')
            dcb.font.bold = True
            dcb.font.color.rgb = BRACKET_COLOR
            rPr_dcb = dcb._element.get_or_add_rPr()
            shd_dcb = OxmlElement('w:shd')
            shd_dcb.set(qn('w:val'), 'clear')
            shd_dcb.set(qn('w:color'), 'auto')
            shd_dcb.set(qn('w:fill'), CREAM_FILL)
            rPr_dcb.append(shd_dcb)

            # Part 2 Receipt options - Para 27, 28, 29
            # Para 27 - first option with opening bracket
            format_option_first(paragraphs[27],
                "furnished to the hospital managers through their internal mail system",
                strike=False)

            # Para 28 - middle option
            format_option_middle(paragraphs[28],
                "furnished to the hospital managers, or a person authorised by them to receive it, by means of electronic communication",
                strike=False)

            # Para 29 - last option (no closing bracket yet - continues to para 31)
            format_option_middle(paragraphs[29],
                "delivered to me in person as someone authorised by the hospital managers to receive this report at [time]",
                strike=False)

            # Para 30 - "on [date]" with cream highlight
            for run in paragraphs[30].runs:
                run.text = ""
            while len(paragraphs[30].runs) > 1:
                paragraphs[30]._element.remove(paragraphs[30].runs[-1]._element)
            pPr30 = paragraphs[30]._element.get_or_add_pPr()
            for old_shd in pPr30.findall(qn('w:shd')):
                pPr30.remove(old_shd)
            if paragraphs[30].runs:
                paragraphs[30].runs[0].text = 'on [date]'
                ct30 = paragraphs[30].runs[0]
            else:
                ct30 = paragraphs[30].add_run('on [date]')
            ct30.font.name = 'Arial'
            ct30.font.size = Pt(12)
            rPr30 = ct30._element.get_or_add_rPr()
            shd30 = OxmlElement('w:shd')
            shd30.set(qn('w:val'), 'clear')
            shd30.set(qn('w:color'), 'auto')
            shd30.set(qn('w:fill'), CREAM_FILL)
            rPr30.append(shd30)

            # Para 31 - placeholder with closing bracket
            for run in paragraphs[31].runs:
                run.text = ""
            while len(paragraphs[31].runs) > 1:
                paragraphs[31]._element.remove(paragraphs[31].runs[-1]._element)
            pPr31 = paragraphs[31]._element.get_or_add_pPr()
            for old_shd in pPr31.findall(qn('w:shd')):
                pPr31.remove(old_shd)
            if paragraphs[31].runs:
                paragraphs[31].runs[0].text = '                                                                        '
                ct31 = paragraphs[31].runs[0]
            else:
                ct31 = paragraphs[31].add_run('                                                                        ')
            ct31.font.name = 'Arial'
            ct31.font.size = Pt(12)
            rPr31 = ct31._element.get_or_add_rPr()
            shd31 = OxmlElement('w:shd')
            shd31.set(qn('w:val'), 'clear')
            shd31.set(qn('w:color'), 'auto')
            shd31.set(qn('w:fill'), CREAM_FILL)
            rPr31.append(shd31)
            cb31 = paragraphs[31].add_run(']')
            cb31.font.name = 'Arial'
            cb31.font.size = Pt(12)
            cb31.font.bold = True
            cb31.font.color.rgb = BRACKET_COLOR
            rPr_cb31 = cb31._element.get_or_add_rPr()
            shd_cb31 = OxmlElement('w:shd')
            shd_cb31.set(qn('w:val'), 'clear')
            shd_cb31.set(qn('w:color'), 'auto')
            shd_cb31.set(qn('w:fill'), CREAM_FILL)
            rPr_cb31.append(shd_cb31)

            # Part 2 Signed - Para 32
            set_labeled_entry(paragraphs[32], "Signed", "", " on behalf of the hospital managers")

            # Part 2 PRINT NAME/Date - Para 33
            for run in paragraphs[33].runs:
                run.text = ""
            while len(paragraphs[33].runs) > 1:
                paragraphs[33]._element.remove(paragraphs[33].runs[-1]._element)
            pPr33 = paragraphs[33]._element.get_or_add_pPr()
            for old_shd in pPr33.findall(qn('w:shd')):
                pPr33.remove(old_shd)
            # PRINT NAME label
            if paragraphs[33].runs:
                paragraphs[33].runs[0].text = 'PRINT NAME'
                pn = paragraphs[33].runs[0]
            else:
                pn = paragraphs[33].add_run('PRINT NAME')
            pn.font.name = 'Arial'
            pn.font.size = Pt(12)
            # PRINT NAME placeholder
            pnob = paragraphs[33].add_run('[')
            pnob.font.bold = True
            pnob.font.color.rgb = BRACKET_COLOR
            rPr_pnob = pnob._element.get_or_add_rPr()
            shd_pnob = OxmlElement('w:shd')
            shd_pnob.set(qn('w:val'), 'clear')
            shd_pnob.set(qn('w:color'), 'auto')
            shd_pnob.set(qn('w:fill'), CREAM_FILL)
            rPr_pnob.append(shd_pnob)
            pnct = paragraphs[33].add_run('                                        ')
            rPr_pnct = pnct._element.get_or_add_rPr()
            shd_pnct = OxmlElement('w:shd')
            shd_pnct.set(qn('w:val'), 'clear')
            shd_pnct.set(qn('w:color'), 'auto')
            shd_pnct.set(qn('w:fill'), CREAM_FILL)
            rPr_pnct.append(shd_pnct)
            pncb = paragraphs[33].add_run(']')
            pncb.font.bold = True
            pncb.font.color.rgb = BRACKET_COLOR
            rPr_pncb = pncb._element.get_or_add_rPr()
            shd_pncb = OxmlElement('w:shd')
            shd_pncb.set(qn('w:val'), 'clear')
            shd_pncb.set(qn('w:color'), 'auto')
            shd_pncb.set(qn('w:fill'), CREAM_FILL)
            rPr_pncb.append(shd_pncb)
            # Date label
            pdl = paragraphs[33].add_run(' Date')
            pdl.font.name = 'Arial'
            pdl.font.size = Pt(12)
            # Date placeholder
            pdob = paragraphs[33].add_run('[')
            pdob.font.bold = True
            pdob.font.color.rgb = BRACKET_COLOR
            rPr_pdob = pdob._element.get_or_add_rPr()
            shd_pdob = OxmlElement('w:shd')
            shd_pdob.set(qn('w:val'), 'clear')
            shd_pdob.set(qn('w:color'), 'auto')
            shd_pdob.set(qn('w:fill'), CREAM_FILL)
            rPr_pdob.append(shd_pdob)
            pdct = paragraphs[33].add_run('                                        ')
            rPr_pdct = pdct._element.get_or_add_rPr()
            shd_pdct = OxmlElement('w:shd')
            shd_pdct.set(qn('w:val'), 'clear')
            shd_pdct.set(qn('w:color'), 'auto')
            shd_pdct.set(qn('w:fill'), CREAM_FILL)
            rPr_pdct.append(shd_pdct)
            pdcb = paragraphs[33].add_run(']')
            pdcb.font.bold = True
            pdcb.font.color.rgb = BRACKET_COLOR
            rPr_pdcb = pdcb._element.get_or_add_rPr()
            shd_pdcb = OxmlElement('w:shd')
            shd_pdcb.set(qn('w:val'), 'clear')
            shd_pdcb.set(qn('w:color'), 'auto')
            shd_pdcb.set(qn('w:fill'), CREAM_FILL)
            rPr_pdcb.append(shd_pdcb)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form H1 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        """Get current form state for saving."""
        gender = "male" if self.gender_male.isChecked() else ("female" if self.gender_female.isChecked() else ("other" if self.gender_other.isChecked() else ""))
        return {
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "prac_name": self.prac_name.text(),
            "dx_primary": self.dx_primary.currentText(),
            "prac_in_charge": self.prac_in_charge.isChecked(),
            "clinician_rmp": self.clinician_rmp.isChecked(),
            "patient_name": self.patient_name.text(),
            "gender": gender,
            "reasons": self.cards["reasons"]["card"].get_content(),
            "delivery_internal": self.delivery_internal.isChecked(),
            "delivery_electronic": self.delivery_electronic.isChecked(),
            "delivery_hand": self.delivery_hand.isChecked(),
            "internal_time": self.internal_time.time().toString("HH:mm"),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
        }

    def set_state(self, state: dict):
        """Restore form state from saved data."""
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        self.prac_name.setText(state.get("prac_name", ""))
        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)

        if state.get("prac_in_charge", True):
            self.prac_in_charge.setChecked(True)
        else:
            self.prac_nominee.setChecked(True)

        if state.get("clinician_rmp", True):
            self.clinician_rmp.setChecked(True)
        else:
            self.clinician_ac.setChecked(True)

        self.patient_name.setText(state.get("patient_name", ""))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        else:
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
        reasons = state.get("reasons", "")
        self._reasons_narrative = reasons
        self.reasons_preview.setText(reasons)
        self.cards["reasons"]["card"].set_content(reasons)

        if state.get("delivery_internal", True):
            self.delivery_internal.setChecked(True)
        elif state.get("delivery_electronic", False):
            self.delivery_electronic.setChecked(True)
        else:
            self.delivery_hand.setChecked(True)

        if state.get("internal_time"):
            self.internal_time.setTime(QTime.fromString(state["internal_time"], "HH:mm"))

        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))

        self._update_all_previews()

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[H1Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[H1Form] Set gender: {gender}")
