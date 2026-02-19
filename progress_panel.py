# ============================================================
# PROGRESS PANEL - Progress Report with Risk Level Timeline
# Styled to match Risk Overview Panel
# ============================================================

from __future__ import annotations

import re
from datetime import datetime, timedelta
from collections import defaultdict
from typing import Dict, List, Any
import io

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QScrollArea,
    QFrame, QSizeGrip, QSizePolicy, QToolTip, QTextBrowser, QSplitter, QComboBox
)
from PySide6.QtCore import Qt, Signal, QPoint, QUrl
from PySide6.QtGui import QPixmap

from patient_history_panel_shared import CollapsibleSection, apply_macos_blur
from timeline_builder import build_timeline
from shared_data_store import get_shared_store
import random

# ============================================================
# PHRASE DICTIONARY - Varied expressions to avoid repetition
# ============================================================

PHRASE_VARIANTS = {
    # Admission phrases (subject + verb)
    "required_admission": [
        "required an admission",
        "was admitted to hospital",
        "needed hospital admission",
        "was hospitalised",
        "required inpatient care",
        "was admitted for inpatient treatment",
        "needed to be admitted",
        "required a period of inpatient care",
        "was taken into hospital",
        "needed hospitalisation",
        "required admission to the ward",
        "was admitted to the unit",
        "needed inpatient admission",
        "required hospital treatment",
        "was brought into hospital",
        "needed to be hospitalised",
        "required a hospital admission",
        "was admitted for treatment",
        "needed admission for stabilisation",
        "required inpatient admission",
    ],

    # Discharge phrases
    "was_discharged": [
        "was discharged",
        "left hospital",
        "was discharged from hospital",
        "returned to the community",
        "was discharged back to the community",
        "left the ward",
        "was discharged from the unit",
        "returned home",
        "was discharged from inpatient care",
        "left inpatient care",
        "was discharged to community care",
        "transitioned back to community",
        "completed the inpatient stay",
        "was discharged from the ward",
        "returned to community services",
        "was discharged into community care",
        "left the unit",
        "was discharged following stabilisation",
        "was discharged home",
        "was discharged with community follow-up",
    ],

    # Duration intro phrases (no "approximately" as duration_desc already has it)
    "duration_intro": [
        "which lasted",
        "spanning",
        "for a period of",
        "lasting",
        "for",
        "extending over",
        "continuing for",
        "over a period of",
        "which extended for",
        "over",
        "which continued for",
        "for a duration of",
        "totalling",
        "amounting to",
        "of",
        "running to",
        "covering",
        "stretching to",
        "coming to",
        "equating to",
    ],

    # Post-discharge intro phrases
    "following_discharge": [
        "Following discharge",
        "After leaving hospital",
        "On return to the community",
        "After discharge",
        "Upon leaving hospital",
        "Once discharged",
        "After being discharged",
        "On leaving the ward",
        "Post-discharge",
        "After returning to the community",
        "Following return to community care",
        "After leaving inpatient care",
        "On discharge",
        "After the inpatient stay",
        "Following the hospital admission",
        "After hospitalisation",
        "On returning home",
        "Following the admission",
        "After leaving the unit",
        "Once back in the community",
    ],

    # Short time in community (before readmission) - use pronoun_cap for sentence start
    "short_community_time": [
        "{pronoun_cap} was only {time} in the community before",
        "{pronoun_cap} remained in the community for just {time} before",
        "{pronoun_cap} spent only {time} at home before",
        "{pronoun_cap} managed only {time} in the community before",
        "Only {time} passed before",
        "{pronoun_cap} was back in the community for just {time} before",
        "After just {time},",
        "{pronoun_cap} lasted only {time} outside hospital before",
        "Within {time},",
        "{pronoun_cap} stayed in the community for only {time} before",
        "Just {time} later,",
        "{pronoun_cap} was only {time} out of hospital before",
        "Barely {time} passed before",
        "{pronoun_cap} was home for just {time} before",
        "After only {time},",
        "{pronoun_cap} coped for only {time} before",
        "{pronoun_cap} remained home for just {time} before",
        "Merely {time} elapsed before",
        "{pronoun_cap} was discharged for only {time} before",
        "A brief {time} passed before",
    ],

    # Readmission phrases
    "requiring_readmission": [
        "requiring readmission",
        "needing to be readmitted",
        "being readmitted",
        "requiring further admission",
        "needing rehospitalisation",
        "being hospitalised again",
        "requiring another admission",
        "needing to return to hospital",
        "being admitted again",
        "requiring return to inpatient care",
        "needing further inpatient treatment",
        "being taken back into hospital",
        "requiring a further period in hospital",
        "needing another hospital stay",
        "being readmitted to the ward",
        "requiring rehospitalisation",
        "needing to be hospitalised again",
        "being brought back into hospital",
        "requiring additional inpatient care",
        "needing return to the ward",
    ],

    # Longer community period
    "stable_community": [
        "{pronoun} remained in the community for {time}, engaging with community services",
        "{pronoun} stayed stable in the community for {time}, supported by the community team",
        "{pronoun} continued in the community for {time} with ongoing support",
        "{pronoun} was maintained in the community for {time} under community care",
        "{pronoun} lived in the community for {time}, receiving regular follow-up",
        "{pronoun} remained at home for {time} with community support",
        "{pronoun} was supported in the community for {time}",
        "{pronoun} continued under community care for {time}",
        "{pronoun} stayed well in the community for {time}",
        "{pronoun} remained stable for {time} in the community",
        "{pronoun} was followed up in the community for {time}",
        "{pronoun} continued at home for {time} with regular contact",
        "{pronoun} was managed in the community for {time}",
        "{pronoun} remained under community services for {time}",
        "{pronoun} lived at home for {time} with ongoing monitoring",
        "{pronoun} stayed in the community for {time}, attending appointments regularly",
        "{pronoun} was maintained at home for {time}",
        "{pronoun} continued well for {time} in the community",
        "{pronoun} remained in community care for {time}",
        "{pronoun} was supported at home for {time}",
    ],

    # Continued engagement (no further admissions)
    "continued_engagement": [
        "{pronoun} continued to engage with community services",
        "{pronoun} remained under community follow-up",
        "{pronoun} continued with ongoing community support",
        "{pronoun} stayed engaged with the community team",
        "{pronoun} remained in contact with services",
        "{pronoun} continued under community care",
        "{pronoun} maintained engagement with the team",
        "{pronoun} remained stable in the community",
        "{pronoun} continued to attend appointments",
        "{pronoun} stayed under community follow-up",
        "{pronoun} continued with regular community contact",
        "{pronoun} remained well supported in the community",
        "{pronoun} continued to be monitored in the community",
        "{pronoun} maintained contact with community services",
        "{pronoun} remained engaged with ongoing care",
        "{pronoun} continued to receive community support",
        "{pronoun} stayed connected with services",
        "{pronoun} continued under the care of the community team",
        "{pronoun} remained in regular contact with the team",
        "{pronoun} continued to be followed up",
    ],

    # Trigger phrases
    "admission_trigger": [
        "This admission was triggered by",
        "The admission followed",
        "This was precipitated by",
        "{pronoun_cap} was admitted following",
        "The hospitalisation was prompted by",
        "This admission came after",
        "The need for admission arose from",
        "{pronoun_cap} required admission due to",
        "This was necessitated by",
        "The admission was precipitated by",
        "Hospitalisation followed",
        "This admission resulted from",
        "{pronoun_cap} was hospitalised after",
        "The need for inpatient care followed",
        "This admission was prompted by",
        "The hospitalisation came after",
        "{pronoun_cap} needed admission following",
        "This was triggered by",
        "The admission arose from",
        "Inpatient care was required after",
    ],

    # Repeated presentation phrases - always list symptoms with qualifier
    "repeated_presentation": [
        "Again, {pronoun} presented with {symptoms}",
        "{pronoun_cap} presented with {symptoms}, as per previous admissions",
        "{pronoun_cap} presented with {symptoms}, a similar picture to previous admissions",
        "{symptoms} recurred, necessitating admission",
        "As before, {pronoun} presented with {symptoms}",
        "{pronoun_cap} again presented with {symptoms}",
        "{symptoms}, symptoms which had been recurrent, led to this admission",
        "{pronoun_cap} presented with a recurrence of {symptoms}",
        "Familiar symptoms of {symptoms} prompted this admission",
        "{pronoun_cap} once again presented with {symptoms}",
        "There was a recurrence of {symptoms}",
        "{symptoms} re-emerged, consistent with previous presentations",
        "{pronoun_cap} relapsed with {symptoms}, mirroring earlier admissions",
        "The recurring pattern of {symptoms} was again evident",
        "{pronoun_cap} presented with {symptoms}, symptoms which had been seen before",
        "A relapse characterised by {symptoms} led to readmission",
        "{pronoun_cap} experienced a return of {symptoms}",
        "{symptoms} recurred, in keeping with {pronoun_poss} established pattern",
        "The now familiar picture of {symptoms} precipitated admission",
        "{pronoun_cap} deteriorated with {symptoms}, as in previous episodes",
    ],

    # During admission phrases
    "during_admission": [
        "During this admission",
        "While in hospital",
        "Throughout the inpatient stay",
        "During the hospitalisation",
        "While on the ward",
        "During this period in hospital",
        "Throughout the admission",
        "While an inpatient",
        "During the hospital stay",
        "While hospitalised",
        "Over the course of the admission",
        "During inpatient treatment",
        "While receiving inpatient care",
        "Throughout hospitalisation",
        "During this inpatient episode",
        "While in the unit",
        "During the ward stay",
        "While admitted",
        "Over the admission period",
        "During this hospital stay",
    ],

    # Improvement phrases (use {pronoun_cap} for sentence start, {pronoun_poss_cap} for possessive)
    # NOTE: Do not include specific factors (medication, therapy) as those are added separately
    "improvement": [
        "{pronoun_cap} showed gradual improvement",
        "{pronoun_poss_cap} mental state stabilised",
        "{pronoun_cap} responded well to treatment",
        "{pronoun_cap} began to improve",
        "{pronoun_poss_cap} condition stabilised",
        "{pronoun_cap} showed signs of recovery",
        "{pronoun_poss_cap} presentation improved",
        "{pronoun_cap} became more settled",
        "{pronoun_cap} showed clinical improvement",
        "{pronoun_poss_cap} symptoms reduced",
        "{pronoun_cap} became less symptomatic",
        "{pronoun_cap} showed positive progress",
        "{pronoun_poss_cap} mental state improved",
        "{pronoun_cap} showed good recovery",
        "{pronoun_cap} improved steadily",
        "{pronoun_poss_cap} behaviour settled",
        "{pronoun_cap} made good progress",
    ],

    # Medication change phrases
    "medication_change": [
        "Medication was changed from {old_med} to {new_med}",
        "{pronoun_poss} medication was adjusted, changing from {old_med} to {new_med}",
        "Treatment was modified with a switch from {old_med} to {new_med}",
        "The medication regime changed from {old_med} to {new_med}",
        "{old_med} was switched to {new_med}",
        "A medication change was made from {old_med} to {new_med}",
        "{pronoun_cap} was changed from {old_med} to {new_med}",
        "The treatment was altered from {old_med} to {new_med}",
        "{old_med} was replaced with {new_med}",
        "Medication adjustment involved changing {old_med} to {new_med}",
    ],

    # Stability phrases
    "remained_stable": [
        "{pronoun} remained relatively stable",
        "{pronoun_poss} presentation was stable",
        "{pronoun} stayed well",
        "{pronoun} continued to be stable",
        "{pronoun} remained well",
        "{pronoun_poss} condition was stable",
        "{pronoun} was generally stable",
        "{pronoun} maintained stability",
        "{pronoun} remained settled",
        "{pronoun_poss} mental state was stable",
        "{pronoun} kept well",
        "{pronoun} continued stable",
        "{pronoun} was in a stable phase",
        "{pronoun_poss} presentation remained settled",
        "{pronoun} stayed relatively well",
        "{pronoun} remained in good mental health",
        "{pronoun} continued without significant concerns",
        "{pronoun} was doing well",
        "{pronoun} remained without major incidents",
        "{pronoun_poss} condition continued stable",
    ],

    # Year deterioration (switching from stable to unstable)
    "year_deterioration": [
        "In {year}, there was some deterioration",
        "{year} saw a decline in {pronoun_poss} presentation",
        "Things became more difficult in {year}",
        "{year} marked a period of increased difficulty",
        "In {year}, {pronoun_poss} condition worsened",
        "{year} brought some deterioration",
        "There was a downturn in {year}",
        "In {year}, stability was lost",
        "{year} saw {pronoun_poss} presentation decline",
        "Difficulties emerged in {year}",
        "{pronoun_poss} condition deteriorated in {year}",
        "In {year}, things became more challenging",
        "{year} was marked by deterioration",
        "A decline was noted in {year}",
        "{year} saw increased instability",
        "In {year}, {pronoun} became less stable",
        "{year} brought new challenges",
        "Problems resurfaced in {year}",
        "In {year}, difficulties returned",
        "{year} marked a period of decline",
    ],

    # Year continued instability
    "year_continued_instability": [
        "{year} continued to show instability",
        "Difficulties persisted through {year}",
        "{year} remained a challenging period",
        "The instability continued into {year}",
        "{year} saw ongoing difficulties",
        "Problems continued through {year}",
        "{year} brought continued challenges",
        "Instability persisted in {year}",
        "{year} was similarly challenging",
        "The difficulties extended into {year}",
        "{year} saw further episodes",
        "Concerns continued in {year}",
        "{year} remained difficult",
        "The pattern continued through {year}",
        "{year} brought ongoing concerns",
        "Difficulties carried on into {year}",
        "{year} saw persistent instability",
        "The challenges continued in {year}",
        "{year} showed ongoing problems",
        "Instability remained evident in {year}",
    ],

    # Community period descriptions (after discharge)
    "community_period_intro": [
        "After leaving the unit, {pronoun} continued well for {time} in the community",
        "Following discharge, {pronoun} remained stable for {time}",
        "On return to the community, {pronoun} did well for {time}",
        "{pronoun_cap} managed in the community for {time}",
        "Back in the community, {pronoun} was stable for {time}",
        "{pronoun_cap} remained well for {time} following discharge",
        "Community living went well for {time}",
        "{pronoun_cap} stayed well for {time} after discharge",
        "The post-discharge period of {time} was stable",
        "{pronoun_cap} coped well for {time} in the community",
        "Following the admission, {pronoun} remained well for {time}",
        "{pronoun_cap} was maintained in the community for {time}",
        "Life in the community continued well for {time}",
        "{pronoun_cap} remained stable for {time} post-discharge",
        "The community period lasted {time}",
        "{pronoun_cap} did well for {time} after leaving hospital",
        "Following discharge, {time} passed without major concern",
        "{pronoun_cap} continued in the community for {time}",
        "Post-discharge, {pronoun} remained well for {time}",
        "{pronoun_cap} was supported in the community for {time}",
    ],
}

# "Again" phrases for repeated situations
AGAIN_VARIANTS = {
    "admission_again": [
        "Again, {pronoun}",
        "Once more, {pronoun}",
        "{pronoun_cap} again",
        "As previously, {pronoun}",
        "Similarly to before, {pronoun}",
        "For the second time, {pronoun}",
        "Repeating the earlier pattern, {pronoun}",
        "As had happened before, {pronoun}",
        "Following a similar pattern, {pronoun}",
        "In a recurrence of earlier difficulties, {pronoun}",
    ],
    "trigger_again": [
        "again presented with",
        "once more experienced",
        "again developed",
        "experienced a recurrence of",
        "again showed signs of",
        "had a return of",
        "again reported",
        "experienced renewed",
        "again described",
        "had a relapse of",
    ],
    "symptom_return": [
        "symptoms returned",
        "difficulties re-emerged",
        "problems resurfaced",
        "presentation deteriorated again",
        "similar difficulties emerged",
        "the pattern repeated",
        "symptoms reappeared",
        "difficulties recurred",
        "problems returned",
        "condition worsened again",
    ],
}

class PhraseTracker:
    """Track used phrases to avoid repetition within a narrative."""

    def __init__(self):
        self.used_phrases = defaultdict(set)  # category -> set of used phrase indices
        self.trigger_history = []  # Track triggers to detect repeats

    def get_phrase(self, category: str, **kwargs) -> str:
        """Get a phrase variant, avoiding recently used ones."""
        variants = PHRASE_VARIANTS.get(category, [category])

        # Find unused variants
        used = self.used_phrases[category]
        available_indices = [i for i in range(len(variants)) if i not in used]

        # If all used, reset and start fresh
        if not available_indices:
            self.used_phrases[category] = set()
            available_indices = list(range(len(variants)))

        # Pick a random available phrase
        idx = random.choice(available_indices)
        self.used_phrases[category].add(idx)
        phrase = variants[idx]

        # Substitute any provided kwargs
        for key, value in kwargs.items():
            phrase = phrase.replace(f"{{{key}}}", str(value))

        return phrase

    def get_again_phrase(self, category: str, **kwargs) -> str:
        """Get an 'again' variant for repeated situations."""
        variants = AGAIN_VARIANTS.get(category, ["Again,"])
        phrase = random.choice(variants)

        for key, value in kwargs.items():
            phrase = phrase.replace(f"{{{key}}}", str(value))

        return phrase

    def record_trigger(self, trigger: str):
        """Record a trigger to track for repeats."""
        self.trigger_history.append(trigger)

    def is_repeated_trigger(self, trigger: str) -> bool:
        """Check if this trigger has occurred before."""
        return trigger in self.trigger_history

    def record_complaints(self, complaints: list):
        """Record presenting complaints for an admission."""
        if not hasattr(self, 'complaint_history'):
            self.complaint_history = []
        self.complaint_history.append(set(complaints))

    def are_complaints_similar(self, complaints: list) -> bool:
        """Check if complaints are similar to previous admissions."""
        if not hasattr(self, 'complaint_history') or not self.complaint_history:
            return False
        current = set(complaints)
        # Check if 2+ complaints match any previous admission
        for prev in self.complaint_history:
            if len(current & prev) >= 2:
                return True
        return False

    def record_concerns(self, concerns: list):
        """Record key concerns for an admission."""
        if not hasattr(self, 'concern_history'):
            self.concern_history = []
        self.concern_history.append(set(concerns))

    def are_concerns_repeated(self, concerns: list) -> bool:
        """Check if concerns are the same as previous."""
        if not hasattr(self, 'concern_history') or not self.concern_history:
            return False
        current = set(concerns)
        for prev in self.concern_history:
            if current == prev or (len(current & prev) >= 1 and len(current) <= 2):
                return True
        return False

    def reset(self):
        """Reset tracker for a new narrative."""
        self.used_phrases = defaultdict(set)
        self.trigger_history = []
        self.complaint_history = []
        self.concern_history = []


# Global phrase tracker instance (reset at start of each narrative)
_phrase_tracker = PhraseTracker()


def get_phrase(category: str, **kwargs) -> str:
    """Get a phrase variant using the global tracker."""
    return _phrase_tracker.get_phrase(category, **kwargs)


def get_again_phrase(category: str, **kwargs) -> str:
    """Get an 'again' phrase variant."""
    return _phrase_tracker.get_again_phrase(category, **kwargs)


def reset_phrase_tracker():
    """Reset the phrase tracker for a new narrative."""
    _phrase_tracker.reset()


# ============================================================
# NARRATIVE REFERENCE TRACKING - For clickable links to source notes
# ============================================================

class NarrativeReferenceTracker:
    """Track references from narrative phrases to source notes."""

    def __init__(self):
        self.references = {}  # ref_id -> {date, matched_text, content_snippet} or {dates: [...], multi: True}
        self._counter = 0

    def add(self, phrase: str, source_date, matched_text: str = "", content_snippet: str = "") -> str:
        """
        Register a phrase with its source information.
        Returns a unique reference ID for the HTML anchor.

        Args:
            phrase: The display text
            source_date: Date of the source note
            matched_text: The exact text that was matched (for highlighting)
            content_snippet: A snippet of the note content (for finding exact note)
        """
        self._counter += 1
        ref_id = f"ref_{self._counter}"
        self.references[ref_id] = {
            "date": source_date,
            "matched": matched_text or phrase,
            "content_snippet": content_snippet,
        }
        return ref_id

    def add_multi(self, phrase: str, entries_list: list, matched_text: str = "") -> str:
        """
        Register a phrase with MULTIPLE source entries (for summary links like "175 episodes").
        Returns a unique reference ID for the HTML anchor.

        Args:
            phrase: The display text
            entries_list: List of dicts with 'date', 'keyword', 'content' keys
            matched_text: The text to match/highlight
        """
        self._counter += 1
        ref_id = f"ref_{self._counter}"
        self.references[ref_id] = {
            "multi": True,
            "matched": matched_text or phrase,
            "entries": entries_list,  # List of {date, keyword, content_snippet}
        }
        return ref_id

    def get(self, ref_id: str) -> dict:
        """Get the reference data for a given ID."""
        return self.references.get(ref_id, {})

    def reset(self):
        """Reset for a new narrative."""
        self.references = {}
        self._counter = 0


# Global reference tracker instance
_ref_tracker = NarrativeReferenceTracker()


def make_link(phrase: str, source_date, matched_text: str = "", content_snippet: str = "") -> str:
    """
    Wrap a phrase in a clickable anchor tag with reference tracking.

    Args:
        phrase: The display text for the link
        source_date: The date of the source note (datetime object)
        matched_text: The exact text to highlight when navigating (defaults to phrase)
        content_snippet: A snippet of note content to identify the exact note

    Returns:
        HTML anchor tag string
    """
    import html
    ref_id = _ref_tracker.add(phrase, source_date, matched_text or phrase, content_snippet)
    safe_phrase = html.escape(phrase)
    return f'<a href="#{ref_id}">{safe_phrase}</a>'


def make_multi_link(phrase: str, incidents: list, matched_text: str = "") -> str:
    """
    Wrap a phrase in a clickable anchor tag that references MULTIPLE entries.
    Used for summary links like "175 episodes of self-harm".

    Args:
        phrase: The display text for the link
        incidents: List of incident dicts with 'entry' containing 'date' and 'content'
        matched_text: The text to highlight

    Returns:
        HTML anchor tag string
    """
    import html
    # Build entries list for tracker
    entries_list = []
    for inc in incidents:
        entry = inc.get('entry', {})
        entries_list.append({
            'date': entry.get('date'),
            'keyword': inc.get('keyword', matched_text),
            'content_snippet': entry.get('content', '')[:200] if entry.get('content') else ''
        })
    ref_id = _ref_tracker.add_multi(phrase, entries_list, matched_text or phrase)
    safe_phrase = html.escape(phrase)
    return f'<a href="#{ref_id}">{safe_phrase}</a>'


def reset_reference_tracker():
    """Reset the reference tracker for a new narrative."""
    _ref_tracker.reset()


def get_reference(ref_id: str) -> dict:
    """Get reference data by ID."""
    return _ref_tracker.get(ref_id)


# ============================================================
# ADMISSION DETAIL EXTRACTION
# ============================================================

ADMISSION_TRIGGER_PATTERNS = [
    # Violence/aggression triggers
    (r'\b(assault|attacked|hit|punch|kick|violent|aggression|aggressive)\b', 'physical aggression'),
    (r'\b(threaten|intimidat|verbal abuse|shouting|screaming)\b', 'threatening behaviour'),
    # Self-harm triggers
    (r'\b(self[- ]?harm|cut|overdose|ligature|suicid|attempt)\b', 'self-harm'),
    (r'\b(took.*overdose|swallowed|ingested)\b', 'overdose'),
    # Psychotic symptoms - specific patterns for detailed extraction
    (r'\b(command.*voice|voice.*tell|voice.*command|voice.*instruct)\b', 'command hallucinations'),
    (r'\b(hear.*voice|auditory.*hallucin|voice.*head|voice.*talk)\b', 'auditory hallucinations'),
    (r'\b(see.*thing|visual.*hallucin|seeing.*people|seeing.*shadow)\b', 'visual hallucinations'),
    (r'\b(paranoi|persecutory|being.*watched|being.*followed|being.*monitored|spy|surveillance)\b', 'paranoid ideation'),
    (r'\b(delusion.*grandeur|grandiose.*delusion|special.*power|chosen|messiah|god)\b', 'grandiose delusions'),
    (r'\b(delusion.*reference|ideas.*reference|tv.*about|radio.*about|people.*talk.*about)\b', 'ideas of reference'),
    (r'\b(thought.*broadcast|thought.*insert|thought.*withdraw|thought.*control|passivity)\b', 'first rank symptoms'),
    (r'\b(believ.*poison|food.*poison|drink.*poison|being.*poison)\b', 'persecutory delusions about poisoning'),
    (r'\b(believ.*plot|conspir|against.*me|out.*to.*get)\b', 'persecutory delusions'),
    (r'\b(psychosis|psychotic|break.*reality|thought.*disorder|disorganis)\b', 'psychotic symptoms'),
    # Mood symptoms
    (r'\b(low mood|depress|suicidal ideation|hopeless|worthless)\b', 'low mood'),
    (r'\b(elat|manic|hypomanic|grandiose|pressure.*speech|flight.*ideas)\b', 'elevated mood'),
    # Behavioural
    (r'\b(non[- ]?complian|stopped.*medication|refusing.*treatment|disengag)\b', 'disengagement from treatment'),
    # Absconding - require strong evidence, not just "left the ward" which could be authorized
    (r'\b(absconded|awol|missing from ward|failed to return from leave|unauthorized absence|escaped from)\b', 'absconding'),
    (r'\b(deteriorat|relaps|decompensate|decline)\b', 'deterioration'),
    # Risk
    (r'\b(risk.*harm|risk.*others|risk.*self|safeguard)\b', 'increased risk'),
    (r'\b(section.*136|police|a&e|crisis)\b', 'crisis presentation'),
]

MEDICATION_PATTERNS = [
    # Antipsychotics
    (r'\b(olanzapine|zyprexa)\b', 'Olanzapine'),
    (r'\b(risperidone|risperdal)\b', 'Risperidone'),
    (r'\b(quetiapine|seroquel)\b', 'Quetiapine'),
    (r'\b(aripiprazole|abilify)\b', 'Aripiprazole'),
    (r'\b(clozapine|clozaril)\b', 'Clozapine'),
    (r'\b(haloperidol|haldol)\b', 'Haloperidol'),
    (r'\b(paliperidone|invega|xeplion|trevicta)\b', 'Paliperidone'),
    (r'\b(amisulpride|solian)\b', 'Amisulpride'),
    (r'\b(zuclopenthixol|clopixol)\b', 'Zuclopenthixol'),
    (r'\b(flupentixol|depixol|fluanxol)\b', 'Flupentixol'),
    (r'\b(pipotiazine|piportil)\b', 'Piportil'),
    (r'\b(chlorpromazine|largactil)\b', 'Chlorpromazine'),
    (r'\b(lurasidone|latuda)\b', 'Lurasidone'),
    (r'\b(cariprazine|reagila)\b', 'Cariprazine'),
    # Antidepressants
    (r'\b(sertraline|lustral|zoloft)\b', 'Sertraline'),
    (r'\b(fluoxetine|prozac)\b', 'Fluoxetine'),
    (r'\b(citalopram|cipramil)\b', 'Citalopram'),
    (r'\b(escitalopram|cipralex)\b', 'Escitalopram'),
    (r'\b(mirtazapine|zispin)\b', 'Mirtazapine'),
    (r'\b(venlafaxine|efexor)\b', 'Venlafaxine'),
    (r'\b(duloxetine|cymbalta)\b', 'Duloxetine'),
    (r'\b(paroxetine|seroxat)\b', 'Paroxetine'),
    (r'\b(trazodone|molipaxin)\b', 'Trazodone'),
    # Mood stabilisers
    (r'\b(lithium|priadel|camcolit)\b', 'Lithium'),
    (r'\b(valproate|depakote|epilim|sodium valproate)\b', 'Valproate'),
    (r'\b(carbamazepine|tegretol)\b', 'Carbamazepine'),
    (r'\b(lamotrigine|lamictal)\b', 'Lamotrigine'),
    # Anxiolytics/Hypnotics
    (r'\b(lorazepam|ativan)\b', 'Lorazepam'),
    (r'\b(diazepam|valium)\b', 'Diazepam'),
    (r'\b(clonazepam|rivotril)\b', 'Clonazepam'),
    (r'\b(zopiclone|zimovane)\b', 'Zopiclone'),
    (r'\b(promethazine|phenergan)\b', 'Promethazine'),
    # Anticholinergics
    (r'\b(procyclidine|kemadrin)\b', 'Procyclidine'),
]

# Substance misuse patterns - with specific substances and amounts where possible
# NOTE: These patterns should be used with context checking to avoid false positives
SUBSTANCE_PATTERNS = [
    (r'\b(cannabis|marijuana|weed|skunk)\b', 'cannabis'),
    (r'\b(cocaine|crack|coke)\b', 'cocaine'),
    (r'\b(heroin|opiates?|opiate misuse)\b', 'heroin'),
    (r'\b(amphetamine|speed(?!\s+up)|meth)\b', 'amphetamines'),
    (r'\b(mdma|ecstasy)\b', 'MDMA'),
    (r'\b(benzodiazepine misuse|benzo abuse|diazepam misuse)\b', 'benzodiazepine misuse'),
    (r'\b(spice|synthetic cannabis|synthetic cannabinoid)\b', 'synthetic cannabinoids'),
    (r'\b(nitrous oxide|nos|laughing gas)\b', 'nitrous oxide'),
    (r'\b(ketamine|ket)\b', 'ketamine'),
    # LSD - "acid" alone is too broad (matches "salicylic acid", "folic acid", etc.)
    # Only match "acid" in drug-use contexts
    (r'\b(lsd|dropped\s+acid|took\s+acid|on\s+acid|using\s+acid|used\s+acid|acid\s+trip)\b', 'LSD'),
    # Alcohol - be specific to avoid "eating and drinking well" false positives
    (r'\b(alcohol use|alcohol misuse|alcohol abuse|drunk|intoxicated|heavy drinking|binge drinking|alcohol depend|alcoholic)\b', 'alcohol'),
    (r'\b(\d+\s*units?\s*(?:of\s+)?alcohol)\b', 'alcohol'),
]

# Amount patterns for substances
SUBSTANCE_AMOUNT_PATTERNS = [
    (r'(\d+)\s*(?:units?|drinks?)\s*(?:per|a|each)?\s*(?:day|week|daily|weekly)', 'units'),
    (r'(\d+)\s*(?:bottles?|cans?)\s*(?:of\s+)?(?:beer|wine|cider|lager)', 'bottles'),
    (r'(\d+)\s*(?:g|grams?|bags?)\s*(?:of\s+)?(?:cannabis|cocaine|heroin)', 'grams'),
    (r'(?:daily|heavy|regular|occasional)\s+(?:cannabis|alcohol|cocaine)', 'frequency'),
    (r'(?:binge|heavy)\s+drinking', 'pattern'),
]


def extract_medication_with_dose(content: str, med_name: str) -> str:
    """Extract medication name with dose and frequency if available."""
    content_lower = content.lower()
    med_lower = med_name.lower()

    # Frequency patterns to look for
    freq_map = {
        'od': 'od', 'o.d': 'od', 'once daily': 'od', 'once a day': 'od', 'daily': 'od',
        'bd': 'bd', 'b.d': 'bd', 'twice daily': 'bd', 'twice a day': 'bd',
        'tds': 'tds', 't.d.s': 'tds', 'three times': 'tds', 'three times daily': 'tds',
        'qds': 'qds', 'q.d.s': 'qds', 'four times': 'qds',
        'nocte': 'nocte', 'at night': 'nocte', 'night': 'nocte',
        'mane': 'mane', 'morning': 'mane', 'in the morning': 'mane',
        'prn': 'prn', 'as required': 'prn', 'when required': 'prn',
        'weekly': 'weekly', 'once weekly': 'weekly',
        'fortnightly': 'fortnightly', 'every 2 weeks': 'fortnightly', 'every two weeks': 'fortnightly',
        'monthly': 'monthly', 'every 4 weeks': 'monthly', 'every four weeks': 'monthly',
        'stat': 'stat',
    }

    # Look for dose patterns with frequency
    dose_freq_patterns = [
        # Medication dose frequency: Olanzapine 10mg od
        rf'\b{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s*(od|bd|tds|qds|prn|nocte|mane|daily|twice daily|weekly|stat)\b',
        # Medication dose at night/morning: Olanzapine 10mg at night
        rf'\b{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s*(at night|in the morning|once daily|twice daily)\b',
        # More flexible: dose and frequency with optional words between
        rf'\b{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s+(?:to be taken\s+)?(od|bd|tds|qds|nocte|mane|prn|daily)\b',
        # Split doses: Clozapine 100mg mane 200mg nocte -> just get total or first dose
        rf'\b{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s+(mane|morning)',
        # Total daily dose pattern: Clozapine total 300mg
        rf'\b{med_lower}\s+(?:total(?:\s+daily)?(?:\s+dose)?)\s*[:\-]?\s*(\d+(?:\.\d+)?)\s*(mg|mcg|g)',
        # Current dose: currently on Clozapine 300mg
        rf'(?:currently|presently|now)\s+(?:on|taking)\s+{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s*(od|bd|tds|qds|nocte|mane|prn|daily)?',
        # On medication dose frequency: on Clozapine 300mg nocte
        rf'\bon\s+{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s*(od|bd|tds|qds|nocte|mane|prn|daily)?',
    ]

    for pattern in dose_freq_patterns:
        match = re.search(pattern, content_lower)
        if match:
            dose = match.group(1)
            unit = match.group(2)
            freq = match.group(3) if match.lastindex >= 3 and match.group(3) else None
            # Normalise unit
            if unit in ['micrograms', 'microgram']:
                unit = 'mcg'
            elif unit in ['milligrams', 'milligram']:
                unit = 'mg'
            # Normalise frequency
            if freq:
                for k, v in freq_map.items():
                    if k in freq:
                        freq = v
                        break
                return f"{med_name} {dose}{unit} {freq}"
            else:
                return f"{med_name} {dose}{unit}"

    # Look for dose patterns without frequency (ordered by specificity)
    dose_patterns = [
        # Direct patterns: medication followed by dose
        rf'\b{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g|micrograms?|milligrams?)\b',
        rf'\b{med_lower}\s*[:\-]\s*(\d+(?:\.\d+)?)\s*(mg|mcg|g)\b',
        rf'\b{med_lower}\s*\((\d+(?:\.\d+)?)\s*(mg|mcg|g)\)',
        rf'\b{med_lower}\s*\[(\d+(?:\.\d+)?)\s*(mg|mcg|g)\]',

        # Dose before medication
        rf'(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s+(?:of\s+)?{med_lower}\b',

        # Started/commenced/prescribed patterns
        rf'(?:started|commenced|prescribed|given|taking|on)\s+{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)',
        rf'{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s+(?:started|commenced|prescribed)',

        # Increased/decreased patterns
        rf'{med_lower}\s+(?:increased|decreased|reduced|titrated)\s+to\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)',

        # Current dose patterns
        rf'{med_lower}\s+(?:currently|now)\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)',

        # Wider context - medication within 20 chars of dose
        rf'{med_lower}.{{0,20}}?(\d+(?:\.\d+)?)\s*(mg|mcg|g)',

        # Even wider context - medication within 40 chars of dose
        rf'{med_lower}.{{0,40}}?(\d+(?:\.\d+)?)\s*(mg|mcg|g)',

        # Dose on same line as medication (line-based search)
        rf'^[^\n]*{med_lower}[^\n]*?(\d+(?:\.\d+)?)\s*(mg|mcg|g)',

        # Structured list entry: "Clozapine - 300mg" or "Clozapine: 300mg"
        rf'{med_lower}\s*[-:]\s*(\d+(?:\.\d+)?)\s*(mg|mcg|g)',

        # Dose followed by medication name: "300mg Clozapine" or "300 mg of Clozapine"
        rf'(\d+(?:\.\d+)?)\s*(mg|mcg|g)\s+(?:of\s+)?{med_lower}',
    ]

    for pattern in dose_patterns:
        match = re.search(pattern, content_lower)
        if match:
            dose = match.group(1)
            unit = match.group(2)
            if unit in ['micrograms', 'microgram']:
                unit = 'mcg'
            elif unit in ['milligrams', 'milligram']:
                unit = 'mg'
            return f"{med_name} {dose}{unit}"

    # Depot/LAI patterns (for injectable antipsychotics)
    depot_patterns = [
        rf'{med_lower}\s+(\d+(?:\.\d+)?)\s*(mg)\s*(?:depot|lai|injection|im)',
        rf'{med_lower}\s+(?:depot|lai|injection)\s+(\d+(?:\.\d+)?)\s*(mg)',
        rf'{med_lower}.{{0,20}}?(\d+(?:\.\d+)?)\s*(mg).{{0,15}}?(?:depot|monthly|fortnightly|weekly)',
    ]

    for pattern in depot_patterns:
        match = re.search(pattern, content_lower)
        if match:
            dose = match.group(1)
            # Check for frequency
            if re.search(r'monthly|every\s*4\s*week', content_lower):
                return f"{med_name} {dose}mg depot monthly"
            elif re.search(r'fortnightly|every\s*2\s*week', content_lower):
                return f"{med_name} {dose}mg depot fortnightly"
            elif re.search(r'weekly', content_lower):
                return f"{med_name} {dose}mg depot weekly"
            return f"{med_name} {dose}mg depot"

    return med_name


def extract_medication_list_from_text(content: str) -> list:
    """
    Extract a full medication list with doses from text that might contain medication lists.
    Looks for structured medication entries, ward round notes, discharge summaries, etc.
    Returns list of medications with doses where found.
    """
    results = []
    content_lower = content.lower()

    # Look for medication list sections
    # Patterns that indicate a medication list follows
    list_headers = [
        r'(?:current\s+)?medications?\s*[:\-]',
        r'(?:on\s+)?discharge\s+medications?\s*[:\-]',
        r'admission\s+medications?\s*[:\-]',
        r'prescribed\s*[:\-]',
        r'drug\s+chart\s*[:\-]',
        r'medication\s+(?:list|review)\s*[:\-]',
        r'(?:regular\s+)?meds?\s*[:\-]',
        r'ttah?\s*[:\-]',  # To Take Home / To Take Away
        r'on\s+admission.*(?:was|were)\s+(?:on|taking)\s*[:\-]?',
        r'(?:currently|presently)\s+(?:on|taking)\s*[:\-]?',
    ]

    # First, try to find and extract from medication list sections
    for header_pattern in list_headers:
        header_match = re.search(header_pattern, content_lower)
        if header_match:
            # Get text after the header (up to 500 chars or next section)
            start_pos = header_match.end()
            remaining = content[start_pos:start_pos + 500]

            # Look for medication entries in this section
            # Pattern: MedicationName dose unit [frequency]
            med_entry_pattern = r'\b([A-Z][a-z]+(?:pine|pam|zole|pril|olol|ine|ide|ate|one|tal|lam|ril|fen)?)\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g|micrograms?)(?:\s+(od|bd|tds|qds|nocte|mane|prn|daily|twice\s*daily|at\s*night))?'
            for match in re.finditer(med_entry_pattern, remaining, re.IGNORECASE):
                med_name = match.group(1)
                dose = match.group(2)
                unit = match.group(3)
                freq = match.group(4) if match.group(4) else ''

                # Normalise unit
                if unit.lower() in ['micrograms', 'microgram']:
                    unit = 'mcg'
                else:
                    unit = unit.lower()

                # Normalise frequency
                if freq:
                    freq = freq.lower().replace(' ', '')
                    if freq in ['daily', 'oncedaily']:
                        freq = 'od'
                    elif freq == 'twicedaily':
                        freq = 'bd'
                    elif freq == 'atnight':
                        freq = 'nocte'

                # Build medication string
                if freq:
                    results.append(f"{med_name} {dose}{unit} {freq}")
                else:
                    results.append(f"{med_name} {dose}{unit}")

    # Also look for line-by-line medication entries (common in structured notes)
    # Pattern: start of line (or after bullet/number) with medication
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        # Skip empty lines or headers
        if not line or len(line) < 5:
            continue

        # Check if line looks like a medication entry
        # E.g., "- Clozapine 300mg nocte" or "1. Olanzapine 10mg od"
        line_med_pattern = r'^(?:[-•*]\s*|\d+[.)\s]+)?([A-Z][a-z]+(?:pine|pam|zole|pril|olol|ine|ide|ate|one|tal|lam|ril|fen)?)\s+(\d+(?:\.\d+)?)\s*(mg|mcg|g)(?:\s+(od|bd|tds|qds|nocte|mane|prn|daily|twice\s*daily|at\s*night|in\s*the\s*morning))?'
        match = re.match(line_med_pattern, line, re.IGNORECASE)
        if match:
            med_name = match.group(1)
            dose = match.group(2)
            unit = match.group(3).lower()
            freq = match.group(4) if match.group(4) else ''

            if freq:
                freq = freq.lower().replace(' ', '').replace('inthemorning', 'mane').replace('atnight', 'nocte')
                if freq in ['daily', 'oncedaily']:
                    freq = 'od'
                elif freq == 'twicedaily':
                    freq = 'bd'

            if freq:
                results.append(f"{med_name} {dose}{unit} {freq}")
            else:
                results.append(f"{med_name} {dose}{unit}")

    return results


def dedupe_medication_list(meds: list) -> list:
    """
    Aggressively deduplicate a list of medications.
    - Groups by base medication name (case-insensitive)
    - For each group, keeps only the most detailed version (longest with dose info)
    - Also removes entries where one is a prefix of another (e.g., "Aripiprazole 10mg" vs "Aripiprazole 10mg od")
    - Handles both string format and dict format {'name': str, 'date': date, 'snippet': str}
    """
    if not meds:
        return []

    meds = list(meds)

    # Normalise to dict format
    normalised = []
    for med in meds:
        if isinstance(med, dict):
            normalised.append(med)
        else:
            normalised.append({'name': med, 'date': None, 'snippet': ''})
    meds = normalised

    # First pass: group by base medication name
    base_names = {}
    for med in meds:
        # Extract base name (letters only at start)
        med_name = med['name']
        base_match = re.match(r'^([A-Za-z]+)', med_name)
        if base_match:
            base = base_match.group(1).lower()
            if base not in base_names:
                base_names[base] = []
            base_names[base].append(med)

    # Second pass: for each group, keep only the best version
    result = []
    for base, variants in base_names.items():
        if len(variants) == 1:
            result.append(variants[0])
            continue

        # Check for dosed versions
        dosed = [v for v in variants if re.search(r'\d+(?:\.\d+)?\s*(mg|mcg|g)\b', v['name'], re.IGNORECASE)]

        if not dosed:
            # No doses - just keep first variant
            result.append(variants[0])
            continue

        # Sort by length (longest first) to find the most detailed
        dosed_sorted = sorted(dosed, key=lambda x: len(x['name']), reverse=True)

        # Remove any medication that is a prefix of another (more detailed) one
        # e.g., "Aripiprazole 10mg" is prefix of "Aripiprazole 10mg od"
        to_keep = []
        for med in dosed_sorted:
            # Check if this med is already covered by a longer one we're keeping
            is_prefix_of_existing = False
            for existing in to_keep:
                if existing['name'].lower().startswith(med['name'].lower()):
                    is_prefix_of_existing = True
                    break
            if not is_prefix_of_existing:
                to_keep.append(med)

        # Keep only the most detailed (longest) version
        if to_keep:
            result.append(to_keep[0])
        elif dosed:
            result.append(dosed_sorted[0])

    return result


def format_meds_as_links(meds: list, limit: int = 4) -> str:
    """Format a list of medication dicts as clickable links joined by commas."""
    if not meds:
        return ""

    deduped = dedupe_medication_list(meds)[:limit]
    links = []
    for med in deduped:
        if isinstance(med, dict):
            med_name = med.get('name', '')
            med_date = med.get('date')
            med_snippet = med.get('snippet', '')
            # Create clickable link
            link = make_link(med_name, med_date, med_name, med_snippet)
            links.append(link)
        else:
            # Fallback for string format (no link)
            links.append(med)

    return ", ".join(sorted(links, key=lambda x: x.lower() if isinstance(x, str) else x))


def extract_substance_details(content: str, note_date=None, content_snippet: str = "") -> List[Dict]:
    """Extract substance misuse details with amounts where available.

    Includes context checking to avoid false positives from:
    - Historical/past use mentions
    - Denials
    - "Eating and drinking well" type phrases
    """
    substances = []
    content_lower = content.lower()

    # Split into sentences for context-aware extraction
    sentences = re.split(r'[.!?\n]', content_lower)

    for pattern, substance in SUBSTANCE_PATTERNS:
        match = re.search(pattern, content_lower, re.IGNORECASE)
        if match:
            # Get context around the match (100 chars before and after)
            start = max(0, match.start() - 100)
            end = min(len(content_lower), match.end() + 100)
            context = content_lower[start:end]

            # Check for denial/negation context
            denial_patterns = [
                r'denied?\s+(?:any\s+)?(?:use\s+of\s+)?',
                r'denies?\s+(?:any\s+)?(?:use\s+of\s+)?',
                r'denied?\s+(?:that\s+)?(?:he|she|they)\s+(?:had\s+been|was|were|has\s+been)\s+(?:using|attempting)',
                r'no\s+(?:current\s+)?(?:use\s+of\s+)?',
                r'not\s+(?:currently\s+)?(?:using|taking)',
                r'stopped\s+(?:using|taking)',
                r'quit\s+(?:using|taking)?',
                r'gave\s+up',
                r'abstinent\s+from',
                r'in\s+the\s+past',
                r'used\s+to',
                r'previously',
                r'history\s+of',  # Historical mention
                r'past\s+(?:use|history)',
                r'former',
                r'ex-',
                r'no\s+longer',
                r'claimed\s+(?:that\s+)?(?:he|she|they)\s+(?:smoked|used|took)',  # "claimed he smoked cannabis in the past"
            ]

            is_denied_or_historical = False
            for denial_pattern in denial_patterns:
                # Increased distance from 30 to 60 chars to catch longer phrases like
                # "Denied that he had been using/attempting to use cannabis"
                if re.search(denial_pattern + r'.{0,60}' + re.escape(match.group(0)), context, re.IGNORECASE):
                    is_denied_or_historical = True
                    break
                if re.search(re.escape(match.group(0)) + r'.{0,30}' + r'in\s+the\s+past', context, re.IGNORECASE):
                    is_denied_or_historical = True
                    break

            if is_denied_or_historical:
                continue

            detail = {'substance': substance, 'amount': None, 'frequency': None,
                      'date': note_date, 'content_snippet': content_snippet or content[:100],
                      'matched': match.group(0)}

            # Find the sentence containing the match for amount/frequency extraction
            match_sentence = ""
            for sent in sentences:
                if match.group(0) in sent:
                    match_sentence = sent
                    break

            # Look for amounts ONLY in the same sentence as the substance
            if substance == 'alcohol':
                # Look for units in same sentence
                unit_match = re.search(r'(\d+)\s*(?:units?)', match_sentence)
                if unit_match:
                    detail['amount'] = f"{unit_match.group(1)} units"
                # Look for frequency in same sentence
                freq_match = re.search(r'(\d+)\s*(?:units?)?\s*(?:per|a|each)\s*(day|week)', match_sentence)
                if freq_match:
                    detail['frequency'] = f"{freq_match.group(1)} per {freq_match.group(2)}"
                # Pattern descriptions in same sentence
                if re.search(r'binge|heavy', match_sentence):
                    detail['frequency'] = 'heavy use'
            elif substance == 'cannabis':
                # Look for frequency in same sentence
                if re.search(r'daily|every\s*day', match_sentence):
                    detail['frequency'] = 'daily'
                elif re.search(r'regular', match_sentence):
                    detail['frequency'] = 'regular'
                elif re.search(r'occasional', match_sentence):
                    detail['frequency'] = 'occasional'
                # Amount in same sentence ONLY
                amount_match = re.search(r'(\d+)\s*(?:spliffs?|joints?|bags?)', match_sentence)
                if amount_match:
                    detail['amount'] = amount_match.group(0)
            else:
                # Generic frequency detection in same sentence
                if re.search(r'daily|every\s*day', match_sentence):
                    detail['frequency'] = 'daily'
                elif re.search(r'regular', match_sentence):
                    detail['frequency'] = 'regular'
                elif re.search(r'occasional', match_sentence):
                    detail['frequency'] = 'occasional'
                elif re.search(r'heavy', match_sentence):
                    detail['frequency'] = 'heavy use'

            substances.append(detail)

    return substances


def is_adverse_reaction_context(text: str, med_name: str) -> bool:
    """Check if medication is mentioned as an adverse reaction, allergy, or contraindication."""
    text_lower = text.lower()
    med_lower = med_name.lower()

    # Find the medication in the text
    med_match = re.search(re.escape(med_lower), text_lower)
    if not med_match:
        return False

    # Get surrounding context (150 chars before, 50 after)
    start = max(0, med_match.start() - 150)
    end = min(len(text_lower), med_match.end() + 50)
    context = text_lower[start:end]

    # Adverse reaction patterns
    adverse_patterns = [
        r'adverse\s+reaction\s+to',
        r'allergic\s+to',
        r'allergy\s+to',
        r'intoleran\w*\s+(?:to|of)',
        r'contraindicated',
        r'cannot\s+(?:take|tolerate|have)',
        r'not\s+(?:to\s+be\s+given|to\s+have|suitable)',
        r'sensitivity\s+to',
        r'hypersensitiv\w*\s+to',
        r'anaphyla\w*\s+to',
        r'reaction\s+to',
        r'side\s+effects?\s+(?:from|to|with)',
        r'bad\s+reaction\s+to',
        r'previous\s+(?:adverse|bad|severe)\s+reaction',
        r'known\s+(?:allergy|sensitivity|intolerance)',
        r'avoid\s+(?:due|because)',
        r'stopped\s+due\s+to\s+(?:side|adverse)',
        r'discontinued\s+due\s+to',
    ]

    for pattern in adverse_patterns:
        if re.search(pattern, context):
            return True

    return False


INCIDENT_SEVERITY_PATTERNS = [
    # High severity - assault on staff requires explicit staff mention
    (r'\b(assault|attack|punch|kick|hit|bite|scratch|spit)\w*\s+(?:at\s+|on\s+|towards?\s+)?(?:a\s+)?(?:staff|nurse|doctor|hca|member of staff)\b', 'assault on staff', 3),
    (r'\b(?:staff|nurse|doctor|hca)\s+(?:was\s+)?(?:assault|attack|punch|kick|hit|bit|bitten|scratch|spit)', 'assault on staff', 3),
    (r'\b(assault.*patient|attack.*patient|hit.*patient)\b', 'assault on patient', 3),
    # General physical aggression (assault without specified target)
    # Note: "hit" alone is too ambiguous (typos like "hit was" for "it was") - require context
    (r'\b(assault|attack|punch|kick|bite|scratch)\b(?!\s*(?:on|at|towards?)\s*(?:staff|patient|nurse|peer))', 'physical aggression', 2),
    (r'\b(?:he|she|they|patient|john|james|david|mr|mrs|ms)\s+hit\b', 'physical aggression', 2),
    (r'\bhit\s+(?:him|her|them|me|staff|peer|patient|wall|door|window)\b', 'physical aggression', 2),
    (r'\b(?:was|got|been)\s+hit\b', 'physical aggression', 2),
    # "jump" alone is too broad - "jump into a taxi" is not self-harm
    # Require context like "jump off/from" or "jumped off/from"
    (r'\b(serious.*self[- ]?harm|ligature|hanging|overdose)\b', 'serious self-harm', 3),
    (r'\bjump(?:ed|ing)?\s+(?:off|from|out\s+of|in\s+front)\b', 'serious self-harm', 3),
    (r'\b(seclusion|secluded)\b', 'seclusion', 3),
    (r'\b(im|intramuscular).*medication\b', 'IM medication', 2),
    (r'\b(rapid tranquil|rt given)\b', 'rapid tranquillisation', 3),
    (r'\b(restrain|physical intervention|c&r)\b', 'physical restraint', 2),
    # Medium severity
    (r'\b(threaten|intimidat|verbal.*threat)\b', 'threats', 2),
    (r'\b(damage|broke|smash|destroy)\b', 'property damage', 2),
    (r'\b(self[- ]?harm|cut|scratch.*self)\b', 'self-harm', 2),
    (r'\b(absconded?|awol|missing from ward|failed to return)\b', 'absconding', 2),
    # Lower severity
    (r'\b(verbal.*aggression|shouting|swearing)\b', 'verbal aggression', 1),
    # Agitation - "disturbed" removed as too ambiguous (could mean "interrupted")
    (r'\b(agitat\w*|highly\s+unsettled|very\s+unsettled|extremely\s+unsettled|increasingly\s+unsettled)\b', 'agitation', 1),
    (r'\b(refus.*medication|non[- ]?complian)\b', 'medication refusal', 1),
]

# Key incident patterns - for highlighting specific notable incidents
# Type names are simple - "required" is added in narrative generation
KEY_INCIDENT_PATTERNS = [
    # Seclusion - must explicitly mention seclusion, NOT "section"
    (r'\b(seclusion|secluded)\b', 'seclusion'),
    # Response team
    (r'\b(response team|emergency response|code.*call|alarm activated|staff alarm)\b', 'the response team being called'),
]

# ============================================================
# COMMUNITY ENGAGEMENT PATTERNS
# ============================================================

COMMUNITY_ENGAGEMENT_PATTERNS = [
    # Psychology/Therapy - specific types first
    (r'\bcbt\b', 'CBT (cognitive behavioural therapy)', 'therapy'),
    (r'\bdbt\b', 'DBT (dialectical behaviour therapy)', 'therapy'),
    (r'\b(mentalisation|mbt)\b', 'MBT (mentalisation-based therapy)', 'therapy'),
    (r'\bschema\s*(therapy)?\b', 'schema therapy', 'therapy'),
    (r'\btrauma.{0,10}therapy\b', 'trauma-focused therapy', 'therapy'),
    (r'\b(emdr)\b', 'EMDR', 'therapy'),
    # "counselled" alone is too broad - in nursing notes it usually means "advised"
    # Only match specific counselling contexts
    (r'\b(counselling\s+(?:session|psycholog|therap)|(?:bereavement|grief|trauma)\s+counselling|(?:seen|referred)\s+(?:by|to)\s+counsell|receiving\s+counselling)\b', 'counselling', 'therapy'),
    (r'\bindividual therapy\b', 'individual therapy', 'therapy'),
    (r'\bgroup therapy\b', 'group therapy', 'therapy'),
    (r'\bfamily therapy\b', 'family therapy', 'therapy'),
    # "Systemic" alone is too broad - matches job titles like "Systemic Clinical Practitioner"
    (r'\bsystemic\s+(?:therapy|session|work|intervention)\b', 'systemic therapy', 'therapy'),
    (r'\b(?:had|attended|engaged in|receiving)\s+systemic\b', 'systemic therapy', 'therapy'),
    (r'\bcouples therapy\b', 'couples therapy', 'therapy'),
    (r'\bart therapy\b', 'art therapy', 'therapy'),
    (r'\bmusic therapy\b', 'music therapy', 'therapy'),
    (r'\boccupational therapy\b', 'occupational therapy', 'therapy'),
    # Psychology - "psychology input" is too vague, require actual session/therapy evidence
    (r'\b(psychology\s+session|psychological\s+therapy|psychologist\s+(?:session|appointment)|(?:seen|met)\s+(?:by|with)\s+psychologist|attended\s+psychology)\b', 'psychology', 'therapy'),
    # Clinics
    (r'\b(depot clinic|injection clinic|clozapine clinic|lithium clinic)\b', 'depot/monitoring clinic', 'clinic'),
    (r'\b(outpatient|opa|outpatient appointment)\b', 'outpatient clinic', 'clinic'),
    # Medical review - must show ACTUAL review happened, not just arranged/requested
    # "for medical review" or "arranged for consultant review" is NOT a completed review
    (r'\b((?:had|attended|seen\s+(?:in|for|at)|reviewed\s+(?:by|in)|completed)\s+(?:consultant|psychiatrist|medical)\s+review|(?:consultant|psychiatrist|doctor)\s+(?:reviewed|saw|seen))\b', 'consultant review', 'clinic'),
    (r'\b(blood|bloods|fbc|u&e|lft|lipids|clozapine level|lithium level|prolactin)\b', 'blood monitoring', 'clinic'),
    # Activities/Groups - require evidence of actual attendance/involvement
    # "Mental Health Resource Centre" in an address/signature is NOT attendance
    (r'\b(?:attend(?:s|ed|ing)?|go(?:es|ing)?|visit(?:s|ed|ing)?|at\s+the|engaged\s+with|linked\s+(?:in\s+)?with)\s+(?:the\s+)?(?:day\s+centre|resource\s+centre|community\s+centre)\b', 'day centre', 'activity'),
    (r'\b(?:day\s+centre|resource\s+centre|community\s+centre)\s+(?:staff|worker|visit|attendance|session|group)\b', 'day centre', 'activity'),
    # Education - "course" alone is too broad (matches "course of medication")
    # Require educational context for "course"
    (r'\b(education|college|training)\b', 'education', 'activity'),
    (r'\b(?:doing|started?|enrolled|attend(?:ing|ed)?|taking)\s+(?:a\s+)?course\b', 'education', 'activity'),
    (r'\bcourse\s+(?:in|at|on)\b', 'education', 'activity'),
    # Employment - must show ACTUAL employment, not just wanting/looking for it
    # "wants to find employment" or "looking for work" is NOT employment
    (r'\b(working\s+(?:at|in|for|as)|started\s+(?:work|a\s+job)|got\s+a\s+job|employed\s+(?:at|by|as)|in\s+(?:paid\s+)?employment|supported\s+employment|voluntary\s+work|volunteering\s+(?:at|with))\b', 'employment', 'activity'),
    (r'\b(social group|activity group|recovery group|peer support)\b', 'group activities', 'activity'),
    (r'\b(gym|swimming|sports|yoga|walking group)\b', 'exercise/physical activity', 'activity'),
    (r'\b(volunteer|volunteering)\b', 'volunteering', 'activity'),
    # Contact people (who contact was with) - require evidence of actual contact
    # "Care Coordinator" in a signature is NOT contact - need verbs showing contact happened
    (r'\b(?:met with|seen by|visited by|spoke to|contact from)\s+(?:the\s+)?(?:care co-?ordinator|cc)\b', 'care coordinator', 'person'),
    (r'\b(?:care co-?ordinator|cc)\s+(?:visit|contact|called|met|saw)\b', 'care coordinator', 'person'),
    (r'\b(?:cpn|community\s+(?:mental\s+health\s+)?nurse)\s+(?:visit|contact|called|saw)\b', 'CPN', 'person'),
    (r'\b(?:met with|seen by|visited by|spoke to)\s+(?:the\s+)?(?:cpn|community nurse)\b', 'CPN', 'person'),
    (r'\b(?:support worker|stw|recovery worker)\s+(?:visit|contact|called|saw)\b', 'support worker', 'person'),
    (r'\b(?:met with|seen by|visited by)\s+(?:the\s+)?(?:support worker|stw)\b', 'support worker', 'person'),
    (r'\b(?:social worker|sw)\s+(?:visit|contact|called|saw)\b', 'social worker', 'person'),
    (r'\b(?:met with|seen by|visited by)\s+(?:the\s+)?(?:social worker|sw)\b', 'social worker', 'person'),
    (r'\b(?:psychiatrist|consultant|responsible clinician)\s+(?:review|saw|visit)\b', 'psychiatrist', 'person'),
    (r'\b(?:seen by|reviewed by|met with)\s+(?:the\s+)?(?:psychiatrist|consultant|rc)\b', 'psychiatrist', 'person'),
    # Contact modes (how contact was made)
    (r'\b(home visit|hv|visited at home|domiciliary)\b', 'home visit', 'mode'),
    (r'\b(telephone|phone|call|rang|spoken to)\b', 'telephone', 'mode'),
    (r'\b(clinic|attended|appointment|opa)\b', 'clinic appointment', 'mode'),
]

COMMUNITY_CRISIS_PATTERNS = [
    # A&E - require evidence of actual attendance, not boilerplate crisis info
    # "you can attend the A&E" is template text, not actual attendance
    (r'\b(?:attended?|went to|presented to|taken to|brought to|seen (?:at|in))\s+(?:the\s+)?(?:a&e|a and e|emergency department|casualty)\b', 'A&E attendance'),
    (r'\b(?:a&e|emergency department|casualty)\s+(?:attendance|presentation|visit)\b', 'A&E attendance'),
    (r'\bed attendance\b', 'A&E attendance'),
    # Crisis services
    (r'\b(htt|home treatment|crisis team|crisis service|crisis visit)\b', 'home treatment team'),
    (r'\b(crisis line|crisis call|samaritans|mental health line)\b', 'crisis line contact'),
    # Respite/alternative
    # Note: "haven" requires space after to avoid matching "haven't"
    (r'\b(crisis house|respite|sanctuary|haven)\s', 'crisis respite'),
]

COMMUNITY_CONCERN_PATTERNS = [
    (r'\b(non[- ]?complian|not taking medication|stopped medication|refusing medication)\b', 'medication non-compliance'),
    (r'\b(disengage|not engaging|missed.*appointment|dna|did not attend)\b', 'disengagement'),
    (r'\b(deteriorat|decline|relaps)\b', 'deterioration'),
    (r'\b(safeguard|vulnerable|exploitation|abuse)\b', 'safeguarding'),
    (r'\b(substance|alcohol|cannabis|drug use|intoxicat)\b', 'substance use'),
    (r'\b(risk.*increas|increas.*risk|elevated risk)\b', 'increased risk'),
]


# Pre-compiled negation pattern for performance
_NEGATION_PATTERN = re.compile(r'''
    \b(?:did|does|do)n?'?t?\s+(?:not\s+)?(?:need|want|require|attend|engage)|
    \bnot\s+(?:need|want|requir|attend|engag|been\s+(?:seen|attended|engaged))|
    \bno\s+(?:need|interest|engagement|job|work|employment|response|answer|longer)|
    \b(?:refused?|declined?|rejected?|disengaged?|withdr[ae]w|unemployed|quit)\b|
    \bfailed\s+to\s+(?:attend|engage|contact|reach|respond)|
    \bdna\b|
    \bdenied?\s+(?:that\s+)?(?:he|she|they|having|using|taking|smoking|drinking|any)|
    \b(?:abstain|clean\s+(?:from|of)|negative\s+(?:uds|drug))|
    \b(?:wants?|would\s+like|looking|searching|trying|hoping)\s+(?:to\s+)?(?:find|get|start)|
    \bstrugg\w+\s+to\s+(?:find|get)|
    \bdifficult\w*\s+(?:to\s+)?(?:find|get)|
    \bunable\s+to\s+(?:find|get|contact|call|reach)|
    \bunbale\s+to\s+(?:contact|call|reach)|
    \bout\s+of\s+work|
    \b(?:what\s+)?work(?:ed|ing|s)?\s+well|
    \bdid\s+not\s+work|
    \bno(?:thing)?\s+(?:\w+\s+)?(?:to\s+report|during|noted|observed)|
    \bnil\s+(?:to\s+report|of\s+note)|
    \b(?:got\s+)?no\s+response|
    \bnot\s+contactable|
    \bcould(?:n'?t)?\s+(?:not\s+)?(?:contact|reach|get)|
    \bawait(?:ing)?\s+(?:contact|response|call)\s+from|
    \barranged\s+(?:for|a)|
    \b(?:for|needs?|requires?|referred\s+for)\s+(?:a\s+)?(?:consultant|medical|psychiatrist|doctor)\s+review|
    \bto\s+(?:have|be\s+seen|arrange)|
    \bstopped\s+(?:going|attending|doing)|
    \bgave\s+up|
    \bused\s+to\b
''', re.IGNORECASE | re.VERBOSE)

# Pre-compiled historical pattern for performance
_HISTORICAL_PATTERN = re.compile(r'''
    \b(?:had|previously|in\s+the\s+past|used\s+to|history\s+of|past\s+history|known\s+history|background\s+of)|
    \b(?:when|during|while)\s+(?:he|she|they|his|her|their|an?\s+)?(?:was|were|in\s+hospital|admission|inpatient)|
    \bprevious\s+(?:use|usage|history|conviction|offence)|
    \bformer\s+(?:use|user)|
    \bin\s+(?:his|her|their)\s+(?:adolescen|youth|teenage|childhood|twenties)|
    \bas\s+(?:a\s+)?(?:teenager|adolescent|youth)|
    \byears?\s+ago|
    \bsentenced\s+(?:to|for)|
    \b(?:convicted|charged|arrested|imprisoned)\s+(?:of|for|with)|
    \bindex\s+offen[cs]e|
    \bserving\s+(?:a\s+)?sentence|
    \bpleaded\s+(?:guilty|not\s+guilty)|
    \bfound\s+guilty|
    \bpending\s+(?:trial|court|sentencing)|
    \b(?:when|whilst|while)\s+in\s+seclusion|
    \bprior\s+to\s+seclusion|
    \b(?:after|following)\s+seclusion|
    \btried\s+to\s+(?:restrain|seclude|section)|
    \backnowledged\s+(?:assaulting|hitting|attacking)|
    \bthen\s*[.,;]|
    \bat\s+that\s+time|
    \bon\s+that\s+occasion|
    \b(?:risk|past|forensic|substance|personal|family)\s+history|
    \bh/?o\s+|
    \bon\s+\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b
''', re.IGNORECASE | re.VERBOSE)

# Pre-compiled email signature pattern
_EMAIL_SIG_PATTERN = re.compile(r'''
    \b(?:office|mobile|tel|fax)\s*:\s*\d|
    @(?:\w+\.)?(?:gov\.uk|nhs\.(?:uk|net))\b|
    \b\d{1,3}(?:st|nd|rd|th)?\s+floor\b|
    \b(?:kind\s+regards|best\s+wishes)\b|
    \b(?:team\s+manager|duty\s+email)\s*:|
    \bwww\.\w+\.(?:uk|nhs|gov)|
    \ben[0-9]\s+[0-9][a-z]{2}\b
''', re.IGNORECASE | re.VERBOSE)


def is_negated_or_historical(text: str, match_obj) -> bool:
    """
    Check if a match is negated, declined, or refers to a past/historical event.
    Returns True if the match should be EXCLUDED.
    """
    match_start = match_obj.start()

    # Get the sentence containing the match
    sent_start = max(0, text.rfind('.', 0, match_start))
    if sent_start > 0:
        sent_start += 1
    sent_end = text.find('.', match_start)
    if sent_end == -1:
        sent_end = len(text)
    sentence = text[sent_start:sent_end]

    # Quick check with pre-compiled patterns
    if _NEGATION_PATTERN.search(sentence):
        return True
    if _HISTORICAL_PATTERN.search(sentence):
        return True

    # Check for section-level historical context
    prefix_text = text[:match_start].lower()
    if re.search(r'(?:risk|past|forensic|substance|personal|family)\s+history', prefix_text[-500:], re.IGNORECASE):
        return True

    # Check email signature (context around match)
    context_start = max(0, match_start - 200)
    context_end = min(len(text), match_start + 200)
    context = text[context_start:context_end]

    if len(_EMAIL_SIG_PATTERN.findall(context)) >= 2:
        return True

    return False


def extract_community_details(notes: List[Dict], start_date, end_date, episodes: List[Dict] = None) -> Dict:
    """Extract detailed information about a community period from the notes.

    Args:
        notes: List of note dictionaries
        start_date: Start of the community period
        end_date: End of the community period
        episodes: List of episode dicts with 'type', 'start', 'end' keys to exclude inpatient periods
    """
    details = {
        'medications': [],  # List of {'name': str, 'date': date, 'snippet': str}
        'psychology': [],  # List of {'type': str, 'date': date, 'snippet': str, 'matched': str}
        'clinics': [],     # List of {'type': str, 'date': date, 'snippet': str, 'matched': str}
        'activities': [],  # List of {'type': str, 'date': date, 'snippet': str, 'matched': str}
        'contact_people': {},  # Who contact was with: {name: {'count': int, 'date': date, 'snippet': str}}
        'contact_modes': {},   # How contact was made: {mode: {'count': int, 'date': date, 'snippet': str}}
        'crisis_events': [],
        'concerns': [],
        'incidents': [],  # Incidents during community period
        'substance_misuse': [],  # Substance misuse mentions
        'absconding': [],  # AWOL/absconding events
        'engagement_level': 'unknown',
        'key_events': [],
    }

    if not notes:
        return details

    # Convert dates
    if hasattr(start_date, 'date'):
        start_d = start_date.date()
    else:
        start_d = start_date

    if end_date:
        if hasattr(end_date, 'date'):
            end_d = end_date.date()
        else:
            end_d = end_date
    else:
        end_d = start_d + timedelta(days=365)  # Default to 1 year

    # Build list of inpatient periods to exclude
    inpatient_periods = []
    if episodes:
        for ep in episodes:
            if ep.get('type') == 'inpatient':
                ep_start = ep.get('start')
                ep_end = ep.get('end')
                if ep_start and ep_end:
                    # Convert to date objects if needed
                    if hasattr(ep_start, 'date'):
                        ep_start = ep_start.date()
                    if hasattr(ep_end, 'date'):
                        ep_end = ep_end.date()
                    inpatient_periods.append((ep_start, ep_end))

    def is_during_inpatient(check_date):
        """Check if a date falls within any inpatient period."""
        for ip_start, ip_end in inpatient_periods:
            if ip_start <= check_date <= ip_end:
                return True
        return False

    contact_count = 0
    therapy_sessions = 0
    clinic_visits = 0

    for note in notes:
        note_date = note.get('date')
        if not note_date:
            continue

        if hasattr(note_date, 'date'):
            note_d = note_date.date()
        else:
            note_d = note_date

        # Only process notes within the community period
        if not (start_d <= note_d <= end_d):
            continue

        # CRITICAL: Skip notes that fall within an inpatient admission
        if is_during_inpatient(note_d):
            continue

        content = (note.get('content') or note.get('text') or '').lower()
        note_type = (note.get('type') or '').lower()

        # Extract medications
        original_content = note.get('content') or note.get('text') or ''
        content_snippet = original_content[:100]
        for pattern, med_name in MEDICATION_PATTERNS:
            if re.search(pattern, content, re.IGNORECASE):
                # Skip medications mentioned in adverse reaction context
                if is_adverse_reaction_context(original_content, med_name):
                    continue
                # Check if this medication already exists
                existing_bases = [m['name'].split()[0].lower() for m in details['medications']]
                if med_name.split()[0].lower() not in existing_bases:
                    details['medications'].append({
                        'name': med_name,
                        'date': note_d,
                        'snippet': content_snippet
                    })

        # Extract engagement activities (now with category)
        # Apply negation/historical check to avoid false positives like "did not need counselling"
        for item in COMMUNITY_ENGAGEMENT_PATTERNS:
            pattern, activity_type, category = item
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                # Skip if negated, declined, or refers to historical/in-hospital event
                if is_negated_or_historical(content, match):
                    continue

                if category == 'therapy':
                    if activity_type not in [p['type'] for p in details['psychology']]:
                        details['psychology'].append({'type': activity_type, 'date': note_d, 'snippet': content_snippet, 'matched': match.group(0)})
                    therapy_sessions += 1
                elif category == 'clinic':
                    if activity_type not in [c['type'] for c in details['clinics']]:
                        details['clinics'].append({'type': activity_type, 'date': note_d, 'snippet': content_snippet, 'matched': match.group(0)})
                    clinic_visits += 1
                elif category == 'activity':
                    existing_activities = [a['type'] if isinstance(a, dict) else a for a in details['activities']]
                    if activity_type not in existing_activities:
                        details['activities'].append({'type': activity_type, 'date': note_d, 'snippet': content_snippet, 'matched': match.group(0)})
                elif category == 'person':
                    # Store first occurrence for linking
                    if activity_type not in details['contact_people']:
                        details['contact_people'][activity_type] = {'count': 0, 'date': note_d, 'snippet': content_snippet}
                    details['contact_people'][activity_type]['count'] += 1
                    contact_count += 1
                elif category == 'mode':
                    if activity_type not in details['contact_modes']:
                        details['contact_modes'][activity_type] = {'count': 0, 'date': note_d, 'snippet': content_snippet}
                    details['contact_modes'][activity_type]['count'] += 1

        # Extract crisis events
        for pattern, crisis_type in COMMUNITY_CRISIS_PATTERNS:
            if re.search(pattern, content, re.IGNORECASE):
                details['crisis_events'].append({
                    'type': crisis_type,
                    'date': note_d,
                    'summary': content[:200]
                })
                break  # One crisis type per note

        # Extract concerns
        for pattern, concern_type in COMMUNITY_CONCERN_PATTERNS:
            if re.search(pattern, content, re.IGNORECASE):
                if concern_type not in [c['type'] for c in details['concerns']]:
                    details['concerns'].append({
                        'type': concern_type,
                        'date': note_d
                    })
                break

        # Extract incidents (using subset of INCIDENT_SEVERITY_PATTERNS)
        community_incident_patterns = [
            (r'\b(assault|attack|punch|kick|hit|bite)\b', 'physical aggression'),
            (r'\b(threaten|intimidat|verbal.*threat)\b', 'threats'),
            (r'\b(self[- ]?harm|cut|overdose)\b', 'self-harm'),
            (r'\b(police|arrest|section 136|136 suite)\b', 'police involvement'),
            (r'\b(damage|broke|smash|destroy)\b', 'property damage'),
        ]
        for pattern, inc_type in community_incident_patterns:
            inc_match = re.search(pattern, content, re.IGNORECASE)
            if inc_match:
                # Check for negation - "denied self harm", "no self harm", etc.
                if is_negated_or_historical(content, inc_match):
                    continue
                details['incidents'].append({
                    'type': inc_type,
                    'date': note_d,
                    'content_snippet': content[:100],
                })
                break  # One incident type per note

        # Extract substance misuse
        substance_patterns = [
            (r'\b(cannabis|marijuana|weed|skunk)\b', 'cannabis'),
            (r'\b(cocaine|crack)\b', 'cocaine'),
            (r'\b(alcohol|drinking|drunk|intoxicated)\b', 'alcohol'),
            (r'\b(heroin|opiate)\b', 'opiates'),
            (r'\b(amphetamine|speed(?!\s+up))\b', 'amphetamines'),
            (r'\b(drug use|substance use|illicit)\b', 'drugs'),
        ]
        for pattern, substance in substance_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                # Check for negation/denial using comprehensive check
                if is_negated_or_historical(content, match):
                    continue
                details['substance_misuse'].append({
                    'type': substance,
                    'date': note_d,
                    'content_snippet': content[:100],
                })
                break

        # Note: AWOL/absconding is an inpatient concept - not extracted for community periods
        # In community, non-attendance is tracked separately as DNA/non-engagement

    # Determine engagement level
    period_days = (end_d - start_d).days or 1
    contacts_per_month = (contact_count / period_days) * 30

    if contacts_per_month >= 4:
        details['engagement_level'] = 'high'
    elif contacts_per_month >= 2:
        details['engagement_level'] = 'moderate'
    elif contacts_per_month >= 0.5:
        details['engagement_level'] = 'low'
    else:
        details['engagement_level'] = 'minimal'

    # Sort medications by name
    details['medications'] = sorted(details['medications'], key=lambda x: x['name'].lower())

    # Convert defaultdicts to regular dicts
    details['contact_people'] = dict(details['contact_people'])
    details['contact_modes'] = dict(details['contact_modes'])

    # Summarise key events
    if details['crisis_events']:
        details['key_events'].extend([f"{e['type']} ({e['date'].strftime('%b %Y')})" for e in details['crisis_events'][:3]])

    return details


def extract_admission_details(notes: List[Dict], admission_date, discharge_date, episodes: List[Dict]) -> Dict:
    """Extract detailed information about an admission from the notes."""
    details = {
        'triggers': [],
        'incidents': [],
        'incident_summary': {},  # Categorised incident counts
        'key_incidents': [],  # Most significant incidents to highlight
        'notable_incidents': [],  # Specific notable incidents (seclusion, 1-1, police, etc.)
        'incident_progression': {},  # Early vs late admission incidents
        'mental_state': [],
        'improvement_factors': [],
        'medications_before': [],  # List of {'name': str, 'date': date, 'snippet': str}
        'medications_during': [],
        'medications_after': [],
        'medication_changes': [],
        # New: Rich admission context from first 2 weeks
        'admission_context': {
            'presenting_complaint': [],
            'mental_state_on_admission': [],
            'legal_status': None,
            'source': None,  # Where they came from (A&E, police, community, etc.)
            'medication_status': None,  # e.g., "off Clozapine for a week", "refusing medication"
            'key_concerns': [],
            'narrative_snippets': [],  # Rich text snippets for narrative
            'substance_misuse': [],  # Substance details with amounts
        },
    }

    if not notes:
        return details

    # Convert dates
    if hasattr(admission_date, 'date'):
        adm_d = admission_date.date()
    else:
        adm_d = admission_date

    if discharge_date:
        if hasattr(discharge_date, 'date'):
            dis_d = discharge_date.date()
        else:
            dis_d = discharge_date
    else:
        dis_d = adm_d + timedelta(days=30)  # Assume 30 days if no discharge

    # Calculate admission midpoint for progression analysis
    admission_length = (dis_d - adm_d).days
    midpoint = adm_d + timedelta(days=admission_length // 2)

    # Time windows
    pre_admission_start = adm_d - timedelta(days=14)  # 2 weeks before
    post_discharge_end = dis_d + timedelta(days=7)  # 1 week after
    early_admission_end = adm_d + timedelta(days=14)  # First 2 weeks of admission for general context
    presenting_complaint_end = adm_d + timedelta(days=3)  # Only first 3 days for presenting complaints
    discharge_window_start = dis_d - timedelta(days=3)  # Last 3 days before discharge for discharge meds

    # Incident tracking for progression
    early_incidents = []
    late_incidents = []

    # Track if explicit "no medication" statement found for admission/discharge
    admission_no_medication = False
    discharge_no_medication = False

    def is_no_medication_context(text):
        """Check if text indicates patient is on no medication."""
        no_med_patterns = [
            r'not\s+on\s+any\s+medication',
            r'no\s+(?:current\s+)?medication',
            r'medication\s*(?:free|nil)',
            r'nil\s+medication',
            r'not\s+(?:currently\s+)?(?:on|taking)\s+(?:any\s+)?medication',
            r'off\s+all\s+medication',
            r'stopped\s+all\s+medication',
            r'no\s+psychiatric\s+medication',
            r'medication\s*:\s*(?:none|nil)',
            r'discharge(?:d)?\s+(?:on\s+)?no\s+medication',
            r'discharge(?:d)?\s+without\s+medication',
        ]
        text_lower = text.lower()
        for pattern in no_med_patterns:
            if re.search(pattern, text_lower):
                return True
        return False

    def is_explicit_discharge_medication(text, med_name):
        """Check if medication is explicitly mentioned as discharge medication."""
        text_lower = text.lower()
        med_lower = med_name.lower()

        # Patterns that explicitly indicate discharge medication
        explicit_patterns = [
            rf'discharge(?:d)?\s+on\s+[^.]*{re.escape(med_lower)}',
            rf'on\s+discharge\s*[,:]\s*[^.]*{re.escape(med_lower)}',
            rf'discharge\s+medication\s*[,:]\s*[^.]*{re.escape(med_lower)}',
            rf'ttO\s*[,:]\s*[^.]*{re.escape(med_lower)}',  # To Take Out
            rf'medication\s+on\s+discharge\s*[,:]\s*[^.]*{re.escape(med_lower)}',
        ]
        for pattern in explicit_patterns:
            if re.search(pattern, text_lower):
                return True
        return False

    def is_explicit_admission_medication(text, med_name):
        """Check if medication is explicitly mentioned as admission medication."""
        text_lower = text.lower()
        med_lower = med_name.lower()

        # Patterns that explicitly indicate admission medication
        explicit_patterns = [
            rf'(?:on\s+)?admission\s+(?:medication|meds?)\s*[,:]\s*[^.]*{re.escape(med_lower)}',
            rf'admitted\s+on\s+[^.]*{re.escape(med_lower)}',
            rf'on\s+admission\s*[,:]\s*[^.]*{re.escape(med_lower)}',
            rf'(?:pre[- ]?admission|prior\s+to\s+admission)\s+[^.]*{re.escape(med_lower)}',
            rf'medication\s+(?:on|at)\s+admission\s*[,:]\s*[^.]*{re.escape(med_lower)}',
            rf'clerking[^.]*{re.escape(med_lower)}',  # Admission clerking notes
        ]
        for pattern in explicit_patterns:
            if re.search(pattern, text_lower):
                return True
        return False

    # Helper to check if a match is negated (e.g., "no voices", "not paranoid", "denies depression")
    def is_negated(text, match_obj):
        """Check if the match is negated or is just a section header."""
        # Get prefix text before the match (up to 150 chars to catch comma-separated lists)
        # e.g., "denied hearing voices, paranoid ideation, and suicidal thoughts"
        start = max(0, match_obj.start() - 150)
        prefix = text[start:match_obj.start()].lower()

        # Get suffix text after the match (up to 50 chars)
        end = min(len(text), match_obj.end() + 50)
        suffix = text[match_obj.end():end].lower()

        # Check if this is a section header/label (e.g., "Suicide:" or "Suicide/Self-Harm:")
        # These are category names, not actual symptoms
        immediate_suffix = text[match_obj.end():match_obj.end() + 20].strip()
        if immediate_suffix.startswith(':') or immediate_suffix.startswith('/') or ' risks:' in immediate_suffix.lower():
            # This is a header like "Suicide:", "Self-Harm/Suicide risks:", etc.
            # Get extended suffix (up to 150 chars) to check for denial patterns after the header
            extended_suffix = text[match_obj.end():min(len(text), match_obj.end() + 150)].lower()
            # Check if "denied" or "no" appears after the colon
            if re.search(r'\bdeni\w*\b', extended_suffix, re.IGNORECASE):
                return True
            if re.search(r'\bno\s+(previous|history|current|thoughts?|ideas?|plans?|intent|attempts?|risk)\b', extended_suffix, re.IGNORECASE):
                return True
            if re.search(r'\bcurrently\s+has\s+no\b', extended_suffix, re.IGNORECASE):
                return True
            if re.search(r'\bhas\s+no\s+(thoughts?|ideas?|plans?|intent)\b', extended_suffix, re.IGNORECASE):
                return True

        # Direct negation patterns (immediately before the match)
        direct_negations = [
            r'\bno\s+$',                    # "no voices"
            r'\bnot\s+$',                   # "not hearing"
            r'\bwithout\s+$',               # "without voices"
            r'\bdeni\w*\s+$',               # "denies voices"
            r'\bnegat\w*\s+$',              # "negative for"
            r'\babsen\w*\s*(of\s+)?$',      # "absent", "absence of"
            r'\bfree\s+of\s+$',             # "free of"
            r'\bnever\s+$',                 # "never"
        ]
        for neg in direct_negations:
            if re.search(neg, prefix, re.IGNORECASE):
                return True

        # Phrase-based negations (within the prefix)
        phrase_negations = [
            r'\bno\s+evidence\s+of\b',           # "no evidence of thought disorder"
            r'\bno\s+signs?\s+of\b',             # "no signs of"
            r'\bno\s+symptoms?\s+of\b',          # "no symptoms of"
            r'\bno\s+thoughts?\b',               # "no thoughts self harm", "no thought of"
            r'\bno\s+ideas?\b',                  # "no ideas of self harm"
            r'\bnot\s+thought\s+disordered\b',   # "Not thought disordered"
            r'\bnot\s+(?:\w+\s+)?disordered\b',  # "not disordered", "not thought disordered"
            r'\bno\s+plans?\b',                  # "no plans to"
            r'\bno\s+intent\b',                  # "no intent to"
            r'\bno\s+intention\b',               # "no intention to"
            r'\bno\s+wish\w*\b',                 # "no wishes", "no wish to"
            r'\bno\s+desire\b',                  # "no desire to"
            r'\bno\s+urge\w*\b',                 # "no urges"
            r'\bnot\s+obviously\b',              # "not obviously responding to voices"
            r'\bnot\s+apparently\b',             # "not apparently"
            r'\bnot\s+currently\b',              # "not currently"
            r'\bnot\s+actively\b',               # "not actively"
            r'\brule[ds]?\s+out\b',              # "ruled out", "rules out"
            r'\bhas\s+not\b',                    # "has not"
            r'\bhave\s+not\b',                   # "have not"
            r'\bdoes\s+not\b',                   # "does not"
            r'\bdid\s+not\b',                    # "did not"
            r'\bwas\s+not\b',                    # "was not"
            r'\bwere\s+not\b',                   # "were not"
            r'\bis\s+not\b',                     # "is not"
            r'\bare\s+not\b',                    # "are not"
            r'\bno\s+\w+\s+(of|to)\b',           # "no evidence of", "no response to"
            r'\bnot\s+\w+ing\s+to\b',            # "not responding to voices"
            r'\bwithout\s+any\b',                # "without any"
            r'\bno\s+formal\b',                  # "no formal thought disorder"
            r'\bno\s+overt\b',                   # "no overt"
            r'\bno\s+apparent\b',                # "no apparent"
            r'\bno\s+obvious\b',                 # "no obvious"
            r'\bno\s+current\b',                 # "no current"
            r'\bno\s+active\b',                  # "no active"
            r'\bno\s+reported\b',                # "no reported"
            r'\bno\s+expressed\b',               # "no expressed"
            r'\bno\s+previous\b',                # "no previous attempts"
            # Nursing notation: "No X to report"
            r'\bno\s+(?:\w+\s+)*(?:to\s+report|during|this\s+shift)\b',  # "No aggression to report"
            r'\bnothing\s+to\s+report\b',
            r'\bnil\s+(?:to\s+report|of\s+note)\b',
        ]
        for neg in phrase_negations:
            if re.search(neg, prefix, re.IGNORECASE):
                return True

        # Check for "not experiencing/having/showing" patterns
        if re.search(r'\b(not|no)\b.{0,20}\b(experiencing|having|reporting|showing|displaying|exhibiting|responding)\b', prefix, re.IGNORECASE):
            return True

        # Check for "denies X" or "denied X" anywhere in prefix
        if re.search(r'\bdeni\w*\b', prefix, re.IGNORECASE):
            return True

        # Sentence-level check: if match is in a sentence that starts with or contains "denied/denies"
        # Find the sentence containing the match
        match_pos = match_obj.start()
        # Find sentence start (look for period, newline, or start of text)
        sent_start = max(0, text.rfind('.', 0, match_pos) + 1)
        if sent_start == 0:
            sent_start = max(0, text.rfind('\n', 0, match_pos) + 1)
        # Find sentence end
        sent_end = text.find('.', match_pos)
        if sent_end == -1:
            sent_end = len(text)
        sentence = text[sent_start:sent_end].lower()

        # Check if sentence contains denial patterns
        if re.search(r'\bdeni\w*\b', sentence, re.IGNORECASE):
            return True
        # Check for "no X, Y, Z" patterns at sentence level
        if re.search(r'^[^.]*\bno\s+(?:\w+\s*,?\s*(?:and\s+)?)*', sentence, re.IGNORECASE):
            if re.search(r'\bno\s+\w', sentence[:50], re.IGNORECASE):
                return True

        return False

    # Helper to check if an incident is about another patient, not the current patient
    def is_about_other_patient(text, match_obj):
        """
        Check if the matched incident is about another patient.
        E.g., "Staff was busy attending to another patient on seclusion"
        E.g., "witnessed an assault on another patient"
        E.g., "another patient was aggressive"
        """
        match_start = match_obj.start()
        match_end = match_obj.end()

        # Get surrounding context (sentence or nearby text)
        # Find sentence boundaries
        sent_start = max(0, text.rfind('.', 0, match_start))
        if sent_start > 0:
            sent_start += 1
        sent_end = text.find('.', match_end)
        if sent_end == -1:
            sent_end = len(text)
        sentence = text[sent_start:sent_end].lower()

        # Patterns that indicate the incident is about another patient
        other_patient_patterns = [
            r'another\s+patient',                    # "another patient on seclusion"
            r'other\s+patient',                      # "other patient"
            r'a\s+fellow\s+patient',                 # "a fellow patient"
            r'neighbouring\s+patient',              # "neighbouring patient"
            r'(?:patient\s+)?(?:in\s+)?(?:the\s+)?(?:next|adjacent)\s+(?:bed|room|bay)',  # "patient in next bed"
            r'patient\s+(?:\w+\s+)?was\s+(?:on\s+)?seclusion',  # "patient X was on seclusion"
            r'attending\s+to\s+(?:another|other)',  # "attending to another patient"
            r'dealing\s+with\s+(?:another|other)',  # "dealing with another patient"
            r'busy\s+with\s+(?:another|other)',     # "busy with another patient"
            r'witnessed\s+(?:an?\s+)?(?:assault|incident|fight)',  # "witnessed an assault"
            r'saw\s+(?:an?\s+)?(?:assault|incident|fight)',  # "saw an assault"
            r'(?:another|other)\s+patient\s+(?:was\s+)?(?:being\s+)?(?:secluded|restrained|aggressive)',
            r'peer(?:\'s)?\s+(?:behaviour|aggression|incident)',  # "peer's aggression"
        ]

        for pattern in other_patient_patterns:
            if re.search(pattern, sentence, re.IGNORECASE):
                return True

        return False

    # PRE-SCAN: Check all notes for "no medication" statements near admission/discharge
    # This must be done BEFORE extracting medications to avoid false positives
    for note in notes:
        note_date = note.get('date')
        if not note_date:
            continue
        if hasattr(note_date, 'date'):
            scan_note_d = note_date.date()
        else:
            scan_note_d = note_date

        note_content = (note.get('content') or note.get('text') or '')
        if is_no_medication_context(note_content):
            # If within first 3 days of admission, mark admission as no medication
            if adm_d <= scan_note_d <= adm_d + timedelta(days=3):
                admission_no_medication = True
            # If within discharge window, mark discharge as no medication
            if discharge_window_start <= scan_note_d <= post_discharge_end:
                discharge_no_medication = True

    # Main note processing loop
    for note in notes:
        note_date = note.get('date')
        if not note_date:
            continue

        if hasattr(note_date, 'date'):
            note_d = note_date.date()
        else:
            note_d = note_date

        content = (note.get('content') or note.get('text') or '').lower()
        note_type = (note.get('type') or '').lower()

        # Extract medications from all relevant periods (with doses where available)
        original_content = note.get('content') or note.get('text') or ''

        # First, try to extract structured medication lists (ward rounds, discharge summaries, etc.)
        # This is more likely to have complete dose and frequency information
        is_med_list_note = bool(re.search(
            r'(?:medications?|meds?|drug chart|prescribed|ttah?|ward round|discharge summar|admission)\s*[:\-]',
            content, re.IGNORECASE
        ))

        # Helper to add medication with source tracking
        def add_med(med_list, med_name, note_date, content_snippet):
            # Check if medication already exists (by base name)
            base_name = med_name.split()[0].lower() if med_name else ""
            existing_bases = [m['name'].split()[0].lower() for m in med_list]
            if base_name not in existing_bases:
                med_list.append({
                    'name': med_name,
                    'date': note_date,
                    'snippet': content_snippet[:100] if content_snippet else ""
                })

        content_snippet = original_content[:100] if original_content else ""

        # Check for "no medication" statements near admission or discharge
        if is_no_medication_context(original_content):
            # If within first 3 days of admission, mark admission as no medication
            if adm_d <= note_d <= adm_d + timedelta(days=3):
                admission_no_medication = True
            # If within discharge window, mark discharge as no medication
            if discharge_window_start <= note_d <= post_discharge_end:
                discharge_no_medication = True

        # MEDICATION EXTRACTION - Be strict about admission/discharge medications
        # Only extract medications that are explicitly stated as being taken at that time

        # For discharge medications (medications_after):
        # ONLY record if explicitly stated as discharge medication AND no "no medication" flag
        if discharge_window_start <= note_d <= post_discharge_end and not discharge_no_medication:
            # Check if this note explicitly mentions discharge medication
            is_discharge_note = bool(re.search(
                r'(discharge\s+medication|discharged?\s+on|on\s+discharge|tto|to\s+take\s+out|medication\s+on\s+discharge)',
                content, re.IGNORECASE
            ))

            if is_discharge_note:
                if is_med_list_note:
                    extracted_meds = extract_medication_list_from_text(original_content)
                    for med_entry in extracted_meds:
                        if not is_adverse_reaction_context(original_content, med_entry):
                            if is_explicit_discharge_medication(original_content, med_entry):
                                add_med(details['medications_after'], med_entry, note_d, content_snippet)

                for pattern, med_name in MEDICATION_PATTERNS:
                    if re.search(pattern, content, re.IGNORECASE):
                        if not is_adverse_reaction_context(original_content, med_name):
                            if is_explicit_discharge_medication(original_content, med_name):
                                med_with_dose = extract_medication_with_dose(original_content, med_name)
                                add_med(details['medications_after'], med_with_dose, note_d, content_snippet)

        # For admission medications (medications_before):
        # Be careful - look for explicit admission medication statements or clerking notes
        elif pre_admission_start <= note_d <= adm_d + timedelta(days=1) and not admission_no_medication:
            is_admission_note = bool(re.search(
                r'(admission\s+medication|admitted\s+on|on\s+admission|clerking|pre[- ]?admission)',
                content, re.IGNORECASE
            ))

            if is_admission_note or is_med_list_note:
                if is_med_list_note:
                    extracted_meds = extract_medication_list_from_text(original_content)
                    for med_entry in extracted_meds:
                        if not is_adverse_reaction_context(original_content, med_entry):
                            add_med(details['medications_before'], med_entry, note_d, content_snippet)

                for pattern, med_name in MEDICATION_PATTERNS:
                    if re.search(pattern, content, re.IGNORECASE):
                        if not is_adverse_reaction_context(original_content, med_name):
                            med_with_dose = extract_medication_with_dose(original_content, med_name)
                            add_med(details['medications_before'], med_with_dose, note_d, content_snippet)

        # For during-admission medications: more permissive but still check adverse reactions
        elif adm_d <= note_d <= dis_d:
            if is_med_list_note:
                extracted_meds = extract_medication_list_from_text(original_content)
                for med_entry in extracted_meds:
                    if not is_adverse_reaction_context(original_content, med_entry):
                        add_med(details['medications_during'], med_entry, note_d, content_snippet)

            for pattern, med_name in MEDICATION_PATTERNS:
                if re.search(pattern, content, re.IGNORECASE):
                    if not is_adverse_reaction_context(original_content, med_name):
                        med_with_dose = extract_medication_with_dose(original_content, med_name)
                        add_med(details['medications_during'], med_with_dose, note_d, content_snippet)

        # Pre-admission notes (look for triggers) - prioritize same-day notes
        if pre_admission_start <= note_d <= adm_d:
            # Weight: same day = 3, day before = 2, earlier = 1
            if note_d == adm_d:
                weight = 3
            elif (adm_d - note_d).days <= 1:
                weight = 2
            else:
                weight = 1

            for pattern, trigger_type in ADMISSION_TRIGGER_PATTERNS:
                match = re.search(pattern, content, re.IGNORECASE)
                if match and not is_negated(content, match):
                    # Store trigger with weight for later sorting
                    existing = [t for t in details['triggers'] if t[0] == trigger_type]
                    if not existing:
                        details['triggers'].append((trigger_type, weight, note_d))
                    elif existing[0][1] < weight:
                        # Replace with higher weight version
                        details['triggers'].remove(existing[0])
                        details['triggers'].append((trigger_type, weight, note_d))

        # ALSO look at first 2 weeks of admission for triggers and context
        # Often the detailed reason for admission is documented after admission
        if adm_d <= note_d <= early_admission_end:
            # High weight for early admission notes explaining the admission
            weight = 3 if (note_d - adm_d).days <= 3 else 2

            for pattern, trigger_type in ADMISSION_TRIGGER_PATTERNS:
                match = re.search(pattern, content, re.IGNORECASE)
                if match and not is_negated(content, match):
                    existing = [t for t in details['triggers'] if t[0] == trigger_type]
                    if not existing:
                        details['triggers'].append((trigger_type, weight, note_d))
                    elif existing[0][1] < weight:
                        details['triggers'].remove(existing[0])
                        details['triggers'].append((trigger_type, weight, note_d))

            # Extract rich admission context
            ctx = details['admission_context']

            # Source of admission (A&E, police, community, etc.)
            if not ctx['source']:
                if re.search(r'\b(a&e|a and e|emergency department|casualty|ed\b|cdU)', content, re.IGNORECASE):
                    ctx['source'] = 'A&E'
                elif re.search(r'\b(section.?136|s136|place of safety|police)', content, re.IGNORECASE):
                    ctx['source'] = 'Section 136 (police)'
                elif re.search(r'\b(cto recall|recalled)', content, re.IGNORECASE):
                    ctx['source'] = 'CTO recall'
                elif re.search(r'\b(community|cmht|home treatment|htt)', content, re.IGNORECASE):
                    ctx['source'] = 'the community'

            # Legal status - store with source info for linking
            # Must check for CURRENT status, not future/desired status
            if not ctx['legal_status']:
                legal_snippet = content[:100]
                content_lower = content.lower()

                # Future/desired markers - don't detect these as current status
                future_markers = [
                    'would like to be', 'wants to be', 'wish to be',
                    'if successful', 'if this is successful', 'may be placed',
                    'recommended', 'for the future', 'plan is', 'considering',
                    'the plan is', 'the plan is that', 'planning for',
                    'upon discharge', 'on discharge', 'after discharge',
                    'continue upon', 'will continue', 'towards discharge',
                    'work towards', 'working towards', 'remain at',
                ]
                has_future_marker = any(m in content_lower for m in future_markers)

                # Section 47/49 (prison transfer) - most specific, check first
                if re.search(r'\b(section\s*47|s\.?47|sec\s*47)', content, re.IGNORECASE):
                    if 'notional' in content_lower and '37' in content_lower:
                        ctx['legal_status'] = {'status': 'Section 47/49 (notional 37)', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 47'}
                    elif re.search(r'\b(section\s*49|s\.?49|\b49\b)', content, re.IGNORECASE):
                        ctx['legal_status'] = {'status': 'Section 47/49', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 47/49'}
                    else:
                        ctx['legal_status'] = {'status': 'Section 47', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 47'}

                # Section 37/41 (court order with restrictions)
                elif re.search(r'\b(section\s*37|s37)\b', content, re.IGNORECASE):
                    if re.search(r'\b(section\s*41|s41)\b', content, re.IGNORECASE):
                        ctx['legal_status'] = {'status': 'Section 37/41', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 37/41'}
                    else:
                        ctx['legal_status'] = {'status': 'Section 37', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 37'}

                # Section 3 (treatment)
                elif re.search(r'\b(section\s*3|s3)\b.*\b(mha|mental health act|detained)\b', content, re.IGNORECASE):
                    ctx['legal_status'] = {'status': 'Section 3', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 3'}
                elif re.search(r'\bdetained\b.*\bsection\s*3\b', content, re.IGNORECASE):
                    ctx['legal_status'] = {'status': 'Section 3', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 3'}

                # Section 2 (assessment)
                elif re.search(r'\b(section\s*2|s2)\b.*\b(mha|mental health act|detained)\b', content, re.IGNORECASE):
                    ctx['legal_status'] = {'status': 'Section 2', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 2'}
                elif re.search(r'\bdetained\b.*\bsection\s*2\b', content, re.IGNORECASE):
                    ctx['legal_status'] = {'status': 'Section 2', 'date': note_d, 'snippet': legal_snippet, 'matched': 'Section 2'}

                # CTO revoked
                elif re.search(r'\bcto\b.*\brevok', content, re.IGNORECASE):
                    ctx['legal_status'] = {'status': 'CTO revoked to Section 3', 'date': note_d, 'snippet': legal_snippet, 'matched': 'CTO revoked'}

                # Informal - only if clearly CURRENT status, not desired/future
                elif re.search(r'\binformal\b', content, re.IGNORECASE) and not has_future_marker:
                    # Require strong current markers
                    informal_current = [
                        'is informal', 'informal patient', 'informal admission',
                        'admitted informal', 'currently informal', 'now informal'
                    ]
                    if any(m in content_lower for m in informal_current):
                        ctx['legal_status'] = {'status': 'informal', 'date': note_d, 'snippet': legal_snippet, 'matched': 'informal'}

            # Medication status on admission
            if not ctx['medication_status']:
                # Check for compliance first - if they took medication, don't flag as non-compliant
                compliance_indicators = re.search(r'\b(reluctantly\s+)?took\s+(his|her|their|the)?\s*medication', content, re.IGNORECASE)
                if not compliance_indicators:
                    # More specific patterns to avoid false positives like "walking off staff...medications"
                    # "off" must be directly before medication-related word, not just within 30 chars
                    med_status_match = re.search(r'\b(off\s+(?:his|her|their|the\s+)?(?:clozapine|medication|meds?)|stopped\s+(?:taking\s+)?(?:his|her|their|the\s+)?(?:clozapine|medication|meds?)|not\s+taking\s+(?:his|her|their|the\s+)?(?:clozapine|medication|meds?)|refusing\s+(?:his|her|their|the\s+)?(?:clozapine|medication|meds?)|refused\s+(?:his|her|their|the\s+)?(?:clozapine|medication|meds?))\b', content, re.IGNORECASE)
                    if med_status_match:
                        off_match = re.search(r'off\s+clozapine\s+for\s+(\w+\s+\w+|\d+\s*\w+)', content, re.IGNORECASE)
                        if off_match:
                            ctx['medication_status'] = {
                                'text': f"off Clozapine for {off_match.group(1)}",
                                'date': note_d,
                                'matched': off_match.group(0),
                                'snippet': content[:100]
                            }
                        elif re.search(r'refus\w*\s+(?:his|her|their|the\s+)?(?:medication|meds?)|not\s+consent\w*\s+(?:to\s+)?(?:medication|meds?)', content, re.IGNORECASE):
                            ctx['medication_status'] = {
                                'text': "refusing medication",
                                'date': note_d,
                                'matched': med_status_match.group(0),
                                'snippet': content[:100]
                            }
                        else:
                            ctx['medication_status'] = {
                                'text': "not taking medication",
                                'date': note_d,
                                'matched': med_status_match.group(0),
                                'snippet': content[:100]
                            }

            # Presenting complaint / mental state on admission
            # ONLY extract from first 3 days of admission - not later notes
            if note_d <= presenting_complaint_end:
                # Store as tuples: (label, matched_text, note_date, content_snippet) for traceability
                # content_snippet is used to identify the exact note when multiple notes share a date
                existing_labels = [c[0] if isinstance(c, tuple) else c for c in ctx['presenting_complaint']]
                # Use first 100 chars of content as snippet for note identification
                content_snippet = content[:100] if content else ""

                # Low mood / depression
                match = re.search(r'\b(low[\s\-]*mood|depressed|feeling[\s\-]*low|feeling[\s\-]*down)\b', content, re.IGNORECASE)
                if match and 'low mood' not in existing_labels and not is_negated(content, match):
                    ctx['presenting_complaint'].append(('low mood', match.group(0), note_d, content_snippet))

                # Hearing voices / hallucinations
                match = re.search(r'\b(hear(?:ing)?[\s\-]*voice|voices|auditory[\s\-]*hallucin\w*)\b', content, re.IGNORECASE)
                if match and 'hearing voices' not in existing_labels and not is_negated(content, match):
                    ctx['presenting_complaint'].append(('hearing voices', match.group(0), note_d, content_snippet))

                # Disorganised thinking
                match = re.search(r'\b(jumbled[\s\-]*thought|thought[\s\-]*disorder|disorganis\w*)\b', content, re.IGNORECASE)
                if match and 'disorganised thinking' not in existing_labels and not is_negated(content, match):
                    ctx['presenting_complaint'].append(('disorganised thinking', match.group(0), note_d, content_snippet))

                # Paranoid ideation
                match = re.search(r'\b(paranoi\w*|persecutor\w*|being[\s\-]*watch\w*|being[\s\-]*follow\w*)\b', content, re.IGNORECASE)
                if match and 'paranoid ideation' not in existing_labels and not is_negated(content, match):
                    ctx['presenting_complaint'].append(('paranoid ideation', match.group(0), note_d, content_snippet))

                # Suicidal ideation / self-harm
                match = re.search(r'\b(suicid\w*|self[\s\-]?harm\w*|harm[\s\-]*self|cut[\s\-]*self|overdose\w*)\b', content, re.IGNORECASE)
                if match and 'suicidal ideation/self-harm' not in existing_labels and not is_negated(content, match):
                    ctx['presenting_complaint'].append(('suicidal ideation/self-harm', match.group(0), note_d, content_snippet))

            # Key concerns mentioned
            # Tighter patterns to avoid false positives like "Unable to upload...given his consent"
            consent_match = re.search(r'\b(unable\s+to\s+consent|cannot\s+consent|not\s+(?:able\s+to\s+)?consent(?:ing)?|lacks?\s+(?:mental\s+)?capacity|without\s+capacity)\b', content, re.IGNORECASE)
            if consent_match:
                existing_concern_texts = [c['text'] if isinstance(c, dict) else c for c in ctx['key_concerns']]
                if 'an inability to consent to treatment' not in existing_concern_texts:
                    ctx['key_concerns'].append({
                        'text': 'an inability to consent to treatment',
                        'date': note_d,
                        'matched': consent_match.group(0),
                        'content_snippet': content[:100]
                    })

            # Extract substance misuse details
            substance_details = extract_substance_details(original_content, note_date=note_d, content_snippet=original_content[:100])
            for sub in substance_details:
                # Check if we already have this substance
                existing = [s for s in ctx['substance_misuse'] if s['substance'] == sub['substance']]
                if not existing:
                    ctx['substance_misuse'].append(sub)
                # NOTE: Do NOT merge amounts from different notes - each note stands alone
                # The amount "7 joints" from one note should not be attributed to a different note's cannabis mention
                elif sub.get('amount') and not existing[0].get('amount'):
                    # Only update if existing has NO amount and new one does
                    # This means we use the FIRST note that has an amount for this substance
                    existing[0]['amount'] = sub['amount']
                    existing[0]['date'] = sub['date']  # Update to note with amount
                    existing[0]['content_snippet'] = sub['content_snippet']
                    existing[0]['matched'] = sub.get('matched', existing[0].get('matched'))

            # Extract narrative snippets from medical/nursing notes in first 3 days
            if (note_d - adm_d).days <= 3 and ('medical' in note_type or 'liaison' in content or 'assessment' in content):
                # Look for sentences that explain the admission
                sentences = re.split(r'(?<=[.!?])\s+', content)
                for sent in sentences:
                    # Look for sentences about why they presented/were admitted
                    if re.search(r'\b(present|admit|recall|revok|referred|brought|came)\b', sent, re.IGNORECASE):
                        if len(sent) > 30 and len(sent) < 200:
                            clean_sent = sent.strip()
                            if clean_sent and clean_sent not in ctx['narrative_snippets']:
                                ctx['narrative_snippets'].append(clean_sent)
                                if len(ctx['narrative_snippets']) >= 3:
                                    break

        # During admission notes (look for incidents with severity)
        if adm_d <= note_d <= dis_d:
            for pattern, incident_type, severity in INCIDENT_SEVERITY_PATTERNS:
                incident_match = re.search(pattern, content, re.IGNORECASE)
                # Skip if incident is about another patient, or is negated/historical (e.g., "denied self harm", "when in seclusion")
                if incident_match and not is_about_other_patient(content, incident_match) and not is_negated_or_historical(content, incident_match):
                    incident_data = {
                        'date': note_d,
                        'type': incident_type,
                        'severity': severity,
                        'note_type': note_type,
                        'summary': content[:300] if content else ''
                    }
                    details['incidents'].append(incident_data)

                    # Track for progression
                    if note_d < midpoint:
                        early_incidents.append(incident_data)
                    else:
                        late_incidents.append(incident_data)

                    # Count by type
                    if incident_type not in details['incident_summary']:
                        details['incident_summary'][incident_type] = 0
                    details['incident_summary'][incident_type] += 1

                    # Track key (high severity) incidents
                    if severity >= 3:
                        details['key_incidents'].append(incident_data)

                    break  # Only count one incident type per note

            # Check for notable incidents (seclusion, response team) - with negation and other-patient check
            for pattern, notable_type in KEY_INCIDENT_PATTERNS:
                notable_match = re.search(pattern, content, re.IGNORECASE)
                if notable_match and not is_negated(content, notable_match) and not is_about_other_patient(content, notable_match):
                    # Extract a brief description from the note
                    # Try to get a sentence containing the key term
                    sentences = re.split(r'[.!?]', content)
                    relevant_sentence = ""
                    for sent in sentences:
                        if re.search(pattern, sent, re.IGNORECASE):
                            relevant_sentence = sent.strip()[:150]
                            break

                    # Try to extract the reason/cause for the incident
                    # IMPORTANT: Only look in the SAME SENTENCE as the incident to avoid false attributions
                    reason = ""

                    # Find the sentence containing the notable incident
                    notable_pos = notable_match.start()
                    reason_sent_start = max(0, content.rfind('.', 0, notable_pos))
                    if reason_sent_start > 0:
                        reason_sent_start += 1
                    reason_sent_end = content.find('.', notable_pos)
                    if reason_sent_end == -1:
                        reason_sent_end = len(content)
                    incident_sentence = content[reason_sent_start:reason_sent_end]

                    # Also include the previous sentence for context (reasons often precede incidents)
                    prev_sent_start = max(0, content.rfind('.', 0, reason_sent_start - 1))
                    if prev_sent_start > 0:
                        prev_sent_start += 1
                    prev_sentence = content[prev_sent_start:reason_sent_start] if reason_sent_start > 0 else ""
                    search_text = prev_sentence + " " + incident_sentence

                    reason_patterns = [
                        (r'(?:following|after|due to|because of|in response to|triggered by|in order to (?:avoid|prevent|manage))\s+(.{10,80}?)(?:\.|,|$)', 'following'),
                        (r'(?:assault|attack|hit|punch|kick|bit|spit)\s+(?:on\s+)?(?:a\s+)?(?:staff|nurse|patient|member)', 'assault'),
                        (r'(?:self[- ]?harm|cut|ligature|overdose|head.?bang)', 'self-harm'),
                        (r'(?:aggression|aggressive|violent|violence|threatening)', 'aggression'),
                        (r'(?:agitat\w*|(?:very|highly|extremely)\s+distress|scream|shout)', 'agitation'),
                        (r'(?:absconded?|awol|missing from ward|failed to return|avoid.*abscond|prevent.*abscond|risk.*abscond)', 'absconding'),
                        (r'(?:refusing|non[- ]?complian|not taking)', 'non-compliance'),
                    ]

                    for reason_pattern, reason_type in reason_patterns:
                        reason_match = re.search(reason_pattern, search_text, re.IGNORECASE)
                        if reason_match:
                            if reason_type == 'following':
                                # Extract the specific reason text
                                reason = reason_match.group(1).strip()
                                # Clean up the reason
                                reason = re.sub(r'\s+', ' ', reason)
                                if len(reason) > 60:
                                    reason = reason[:60].rsplit(' ', 1)[0]
                                # Validate: don't use reasons that are clearly unrelated community events
                                # (e.g., "taken mother's car" is not a hospital incident reason)
                                invalid_community_reasons = [
                                    r'\b(car|vehicle|driving|police.*picked|arrested)\b',
                                    r'\b(at home|in the community|outside hospital)\b',
                                ]
                                if any(re.search(pat, reason, re.IGNORECASE) for pat in invalid_community_reasons):
                                    reason = ""  # Reject this reason
                                    continue  # Try next pattern
                            else:
                                reason = reason_type
                            break

                    notable_data = {
                        'date': note_d,
                        'type': notable_type,
                        'description': relevant_sentence if relevant_sentence else content[:100],
                        'reason': reason,
                        'note_type': note_type,
                        'matched': notable_match.group(0),  # Store the matched text for highlighting
                        'content_snippet': content[:100],  # Store content snippet for note identification
                    }

                    # Avoid duplicates on same date with same type
                    if not any(n['date'] == note_d and n['type'] == notable_type for n in details['notable_incidents']):
                        details['notable_incidents'].append(notable_data)
                    break  # Only one notable type per note

            # Mental state observations
            if 'mental state' in note_type or 'mse' in note_type or re.search(r'\b(mental state|mse|presentation)\b', content):
                details['mental_state'].append({
                    'date': note_d,
                    'summary': content[:300] if content else ''
                })

        # Look for improvement factors near discharge (causes: medication, psychology, nursing)
        if dis_d - timedelta(days=14) <= note_d <= dis_d:
            if re.search(r'\b(improv|better|settled|stable|recover|respond.*treatment|ready.*discharge)\b', content, re.IGNORECASE):
                existing_factors = [f['text'] if isinstance(f, dict) else f for f in details['improvement_factors']]
                if re.search(r'\b(medication|medic|tablet|depot|antipsychotic|clozapine)\b', content):
                    if 'medication changes' not in existing_factors:
                        details['improvement_factors'].append({'text': 'medication changes', 'date': note_d, 'snippet': content[:100]})
                # Be specific about psychological therapy - exclude occupational therapy
                # Must mention specific therapy types, not just "therapy" (which could be OT)
                # AND must be actual engagement, not just suggestions/requests for therapy
                # Note: "counselled" alone is excluded - in nursing notes it means "advised", not therapy
                # Only match actual psychological therapy terms - "psychology input" is too vague
                psych_match = re.search(r'\b(psychology\s+session|psychological\s+therapy|psychologist|cbt|dbt|mentalisation|mbt|schema|emdr|individual\s+therapy|group\s+therapy|family\s+therapy|systemic\s+therapy|talking\s+therap|counselling\s+(?:session|psycholog|therap)|(?:bereavement|grief|trauma)\s+counselling|receiving\s+counselling)', content)
                if psych_match and not re.search(r'occupational\s+therap', content, re.IGNORECASE):
                    # Check it's not just a suggestion/request/opinion about needing therapy
                    # Get the sentence containing the match
                    psych_pos = psych_match.start()
                    psych_sent_start = max(0, content.rfind('.', 0, psych_pos) + 1)
                    psych_sent_end = content.find('.', psych_pos)
                    if psych_sent_end == -1:
                        psych_sent_end = len(content)
                    psych_sentence = content[psych_sent_start:psych_sent_end].lower()

                    # Exclude if it's a suggestion/need/request, not actual therapy
                    suggestion_patterns = [
                        r'\bneeds?\s+(?:\w+\s+)?(?:cbt|dbt|therap|psycholog|counselling)',
                        r'\bwants?\s+(?:\w+\s+)?(?:cbt|dbt|therap|psycholog|counselling)',
                        r'\bwould\s+benefit\s+from',
                        r'\bshould\s+(?:have|receive|get)',
                        r'\brecommend\w*\s+(?:\w+\s+)?(?:cbt|dbt|therap|psycholog|counselling)',
                        r'\bbelieves?\s+.*(?:needs?|wants?|should|require)',
                        r'\bfeels?\s+.*(?:needs?|wants?|should|require)',
                        r'\bmother\s+.*(?:needs?|wants?|should)',
                        r'\bfather\s+.*(?:needs?|wants?|should)',
                        r'\bfamily\s+.*(?:needs?|wants?|should)',
                        r'\bmay\s+benefit',
                        r'\bcould\s+benefit',
                        r'\brequires?\s+(?:\w+\s+)?(?:cbt|dbt|therap|psycholog|counselling)',
                    ]
                    is_suggestion = any(re.search(pat, psych_sentence, re.IGNORECASE) for pat in suggestion_patterns)

                    if not is_suggestion and 'psychological intervention' not in existing_factors:
                        details['improvement_factors'].append({'text': 'psychological intervention', 'date': note_d, 'snippet': content[:100]})
                if re.search(r'\b(nurs|staff support|ward routine|structure)\b', content):
                    if 'nursing care' not in existing_factors:
                        details['improvement_factors'].append({'text': 'nursing care', 'date': note_d, 'snippet': content[:100]})
                if re.search(r'\b(insight|engag|complian|accept)\b', content):
                    if 'improved insight' not in existing_factors:
                        details['improvement_factors'].append({'text': 'improved insight', 'date': note_d, 'snippet': content[:100]})

            # Track successful leave separately - it's an OUTCOME not a cause
            leave_match = re.search(r'\b(successful.*leave|leave.*successful|ground.*leave|escorted.*leave|unescorted.*leave|overnight.*leave|community.*leave)\b', content, re.IGNORECASE)
            if leave_match:
                if 'successful_leave' not in details:
                    details['successful_leave'] = {'date': note_d, 'matched': leave_match.group(0), 'snippet': content[:100]}

    # Analyse incident progression
    details['incident_progression'] = {
        'early_count': len(early_incidents),
        'late_count': len(late_incidents),
        'early_severe': len([i for i in early_incidents if i['severity'] >= 3]),
        'late_severe': len([i for i in late_incidents if i['severity'] >= 3]),
        'pattern': 'unknown'
    }

    # Determine pattern
    if len(early_incidents) > len(late_incidents) * 1.5:
        details['incident_progression']['pattern'] = 'improving'
    elif len(late_incidents) > len(early_incidents) * 1.5:
        details['incident_progression']['pattern'] = 'worsening'
    elif len(early_incidents) == 0 and len(late_incidents) == 0:
        details['incident_progression']['pattern'] = 'stable'
    else:
        details['incident_progression']['pattern'] = 'fluctuating'

    # Helper to deduplicate medications - remove plain names if we have doses
    # Define this BEFORE using it for medication changes
    def dedupe_meds(med_list):
        """Deduplicate medications, keeping source info from best entry."""
        if not med_list:
            return []

        # Group by base medication name (case-insensitive)
        base_names = {}
        for med_entry in med_list:
            # Handle both old format (string) and new format (dict)
            if isinstance(med_entry, dict):
                med_name = med_entry.get('name', '')
            else:
                med_name = med_entry
                med_entry = {'name': med_name, 'date': None, 'snippet': ''}

            # Extract base name (everything before the dose number)
            base_match = re.match(r'^([A-Za-z]+)', med_name)
            if base_match:
                base = base_match.group(1).lower()  # Case-insensitive grouping
                if base not in base_names:
                    base_names[base] = []
                base_names[base].append(med_entry)

        # For each base name, keep only the BEST dosed version
        result = []
        for base, variants in base_names.items():
            # Check if any variant has a dose
            dosed = [v for v in variants if re.search(r'\d+(?:\.\d+)?\s*(mg|mcg|g)\b', v['name'], re.IGNORECASE)]
            if dosed:
                # Keep only the most detailed dosed version (longest string = most info)
                best = max(dosed, key=lambda x: len(x['name']))
                result.append(best)
            else:
                # Keep just one plain name
                result.append(variants[0])
        return result

    # Deduplicate medications
    details['medications_before'] = dedupe_meds(details['medications_before'])
    details['medications_during'] = dedupe_meds(details['medications_during'])
    details['medications_after'] = dedupe_meds(details['medications_after'])

    # Identify medication changes (using deduplicated lists)
    meds_before_names = set(m['name'] for m in details['medications_before'])
    meds_during_names = set(m['name'] for m in details['medications_during'])

    new_med_names = meds_during_names - meds_before_names
    # Get the full med entries for new medications
    new_meds = [m for m in details['medications_during'] if m['name'] in new_med_names]

    # Apply deduplication to medication changes (only track started medications, not stopped)
    if new_meds:
        details['medication_changes'].append(('started', dedupe_meds(new_meds)))

    # Sort triggers by weight (highest first) and extract just the names
    # Triggers are stored as (trigger_type, weight, date) tuples
    if details['triggers'] and isinstance(details['triggers'][0], tuple):
        sorted_triggers = sorted(details['triggers'], key=lambda x: -x[1])  # Sort by weight descending
        details['triggers'] = [t[0] for t in sorted_triggers]  # Extract just the trigger names
    details['triggers'] = list(dict.fromkeys(details['triggers']))  # Remove duplicates preserving order
    # Deduplicate improvement_factors (handles both dict and string formats)
    seen_factors = set()
    unique_factors = []
    for f in details['improvement_factors']:
        factor_text = f['text'] if isinstance(f, dict) else f
        if factor_text not in seen_factors:
            seen_factors.add(factor_text)
            unique_factors.append(f)
    details['improvement_factors'] = unique_factors

    # Sort key incidents by date
    details['key_incidents'] = sorted(details['key_incidents'], key=lambda x: x['date'])[:5]  # Top 5

    # Sort notable incidents by date and keep the most significant ones
    # Prioritise seclusion and response team (police removed due to false positives)
    priority_types = ['seclusion', 'the response team being called']
    priority_notable = [n for n in details['notable_incidents'] if n['type'] in priority_types]
    other_notable = [n for n in details['notable_incidents'] if n['type'] not in priority_types]
    details['notable_incidents'] = sorted(priority_notable, key=lambda x: x['date'])[:3]
    if len(details['notable_incidents']) < 3:
        details['notable_incidents'].extend(sorted(other_notable, key=lambda x: x['date'])[:3 - len(details['notable_incidents'])])

    return details


# ============================================================
# TENTPOLE EVENTS - Key milestones in patient journey
# ============================================================

TENTPOLE_PATTERNS = {
    "Tribunal": {
        "color": "#1565c0",
        "patterns": [r'\b(tribunal|mental\s+health\s+tribunal|mht)\b'],
    },
    "Managers Hearing": {
        "color": "#0d47a1",
        "patterns": [r"manager'?s?\s+hearing", r'hospital\s+manager.{0,10}(hearing|review)'],
    },
    "CPA Review": {
        "color": "#2e7d32",
        "patterns": [
            r'\bcpa\s+(review|meeting|held|took\s+place|scheduled|arranged)\b',
            r'\b(cpa|care\s+programme\s+approach)\s+(review|meeting)\b',
            r'\breview.{0,15}(cpa|care\s+programme)\b',
            r'\b(mdt|multi\s*disciplinary)\s+(review|meeting)\s+(held|took\s+place)\b',
            r'\battended\s+(cpa|mdt)\b',
        ],
    },
    "Ward Round": {
        "color": "#388e3c",
        "patterns": [r'\bward\s+round\b'],
    },
    "Ground Leave": {
        "color": "#4caf50",
        "patterns": [r'ground\w?\s+leave\s+(granted|approved|started|used)', r'(first|initial)\s+ground\w?\s+leave'],
    },
    "Escorted Leave": {
        "color": "#8bc34a",
        "patterns": [r'escorted\s+leave\s+(granted|approved|started|used)'],
    },
    "Unescorted Leave": {
        "color": "#cddc39",
        "patterns": [r'unescorted\s+leave\s+(granted|approved|started|used)'],
    },
    "Overnight Leave": {
        "color": "#ffeb3b",
        "patterns": [r'overnight\s+leave', r'overnight\s+stay'],
    },
    "Community Leave": {
        "color": "#ffc107",
        "patterns": [r'community\s+(leave|placement)', r'(section\s+17|s17)\s+leave'],
    },
    "Section Change": {
        "color": "#ff9800",
        "patterns": [r'(section|s)\s*\d+\s+(to|changed)', r'detained\s+under\s+(section|s)\s*\d+'],
    },
    "Medication Change": {
        "color": "#9c27b0",
        "patterns": [r'(started|commenced|initiated)\s+(on\s+)?\w+\s*\d+\s*mg', r'depot\s+(changed|started|increased)'],
    },
}

# Risk patterns for scoring
# All patterns use word boundaries (\b) to match whole words only
RISK_PATTERNS = {
    "Physical Violence": {
        "weight": 10,
        "patterns": [
            r'\b(punched|punching|kicked|kicking|hit|hits|hitting|slapped|slapping|struck|attacked|attacking)\b.{0,15}\b(staff|peer|patient|nurse|hca)\b',
            r'\bassaulted\b.{0,15}\b(staff|peer|patient|nurse)\b',
            r'\bphysically\s+aggressive\b',
            r'\bphysical\s+aggression\b',
            r'\b(restraint|restrained)\s+(required|needed|used|applied)\b',
            r'\brequired\s+(restraint|physical\s+intervention)\b',
        ],
    },
    "Verbal Aggression": {
        "weight": 3,
        "patterns": [
            r'\bverbally\s+(abusive|aggressive)\b',
            r'\bverbal\s+aggression\b',
            r'\bshouting\s+at\s+(staff|peers|patients|nurses)\b',
            r'\b(threatening|threatened)\s+(staff|behaviour|language|towards)\b',
            r'\bmade\s+(threats|a\s+threat)\b',
        ],
    },
    "Self-Harm": {
        "weight": 8,
        "patterns": [
            r'\bself[\s-]?harmed\b',
            r'\bself[\s-]?harming\b',
            r'\battempted\s+self[\s-]?harm\b',
            r'\bepisode\s+of\s+self[\s-]?harm\b',
            r'\bligature\s+(found|made|attempted|used)\b',
            r'\b(took|taken)\s+(an\s+)?overdose\b',
            r'\boverdosed\b',
            r'\b(cut|cutting|cuts)\s+(himself|herself|his|her)\s+(arm|wrist|leg)\b',
            r'\bsuicidal\s+(ideation|thoughts|intent)\b',
            r'\bexpressed\s+.{0,20}\b(kill|end)\s+(himself|herself|my\s+life)\b',
        ],
    },
}

# Negative context patterns - if these appear near a match, it's likely a negation
NEGATIVE_CONTEXT_PATTERNS = [
    r'\b(no|nil|none|denies?|denied|without)\b.{0,15}$',  # Before match
    r'^.{0,15}\b(no|nil|none|denied|not\s+noted|not\s+reported|not\s+present)\b',  # After match
    r'\bno\s+(evidence|history|signs?|indication|report)\s+of\b',
    r'\b(has\s+not|did\s+not|does\s+not|wasn\'t|weren\'t|isn\'t|aren\'t)\b',
    r'\brisk\s+(assessment|factors?)\b.{0,30}$',  # Likely discussing risk assessment, not actual incident
]


def _has_negative_context(text: str, match_start: int, match_end: int) -> bool:
    """Check if a match has negative context (nil, no, denied, etc.)."""
    # Get context around the match (50 chars before and after)
    context_start = max(0, match_start - 50)
    context_end = min(len(text), match_end + 50)

    before_text = text[context_start:match_start].lower()
    after_text = text[match_end:context_end].lower()
    full_context = text[context_start:context_end].lower()

    # Check for negative patterns before the match
    negative_before = [
        r'\b(no|nil|none|denies?|denied|without|lacks?)\s*$',
        r'\b(no|nil|none|denies?|denied|without|lacks?)\s+(any|all|the|a)?\s*$',
        r'\b(no\s+evidence|no\s+history|no\s+signs?|no\s+indication)\s+of\s*$',
        r'\b(has\s+not|did\s+not|does\s+not|hasn\'t|didn\'t|doesn\'t)\s*$',
        r'\b(not\s+noted|not\s+reported|not\s+observed)\s*$',
        r'\bdenied\s+(any|all)?\s*(thoughts?\s+of|ideation\s+of|intent\s+to|intention\s+to)?\s*$',
        # Additional patterns for common clinical negations
        r'\b(not|never)\s+express(ing|ed)?\s*(any)?\s*$',
        r'\bno\s+(current|recent|active|new)?\s*(thoughts?|episodes?|ideation|urges?|intent)\s+(of|to)\s*$',
        r'\b(doesn\'t|does\s+not|don\'t|do\s+not)\s+have\s+(\w+\s+)*$',  # "doesn't have current active"
        r'\b(there\s+)?(were|was|are|is)\s+no\s+(episode|evidence|history|indication)s?\s+(of)?\s*$',
        r'\bno\s+\w+\s+of\s*$',  # "no episode of", "no thoughts of", etc.
        r'\bwith\s+no\s*$',  # "with no self-harming"
        # Handle "or" clauses - negation earlier in phrase carries through
        r'\bor\s+(urges?\s+to|thoughts?\s+of|intent\s+to)?\s*$',
        # "was not X", "were not X" patterns
        r'\b(was|were|is|are)\s+not\s*$',
    ]

    for pattern in negative_before:
        if re.search(pattern, before_text):
            return True

    # Check if negation appears earlier in the broader context (handles "X or Y" constructs)
    broader_negation = [
        r'\b(no|not|nil|none|denied|denies|without)\b.{0,40}\bor\b',
        r'\b(doesn\'t|does\s+not|don\'t)\s+have\b',
    ]
    for pattern in broader_negation:
        if re.search(pattern, before_text):
            return True

    # Check for negative patterns after the match
    negative_after = [
        r'^\s*(nil|none|denied|not\s+noted|not\s+reported)\b',
        r'^\s*-\s*(nil|no|none|denied)\b',
        r'^\s*(were|was|are|is)?\s*(not\s+present|not\s+identified|not\s+expressed|absent)\b',
        r'^\s*(were|was|are|is)?\s*not\s+(noted|reported|observed|evident|identified)\b',
        # Allow for intervening words like "thoughts", "ideation" before negative phrase
        r'^\s*\w*\s*(were|was|are|is)?\s*(not\s+present|absent)\b',
        r'^\s*(thoughts?|ideation)?\s*(were|was|are|is)?\s*(not\s+present|absent|not\s+expressed|denied)\b',
        # "X was not required/needed" patterns
        r'^\s*(was|were|is|are)?\s*not\s+(required|needed|necessary|indicated)\b',
    ]

    for pattern in negative_after:
        if re.search(pattern, after_text):
            return True

    # Check full context for assessment/documentation language
    assessment_patterns = [
        r'\brisk\s+(assessment|screen|factor)',
        r'\b(asked|enquired|assessed)\s+about\b',
        r'\bqueried\s+(re|regarding|about)\b',
    ]

    for pattern in assessment_patterns:
        if re.search(pattern, full_context):
            # Only exclude if also has negative indicator
            if re.search(r'\b(no|nil|denied|denies|negative)\b', full_context):
                return True

    return False


def _normalise_date(d):
    """Convert any date format to datetime."""
    if d is None:
        return None
    if isinstance(d, datetime):
        return d
    try:
        import pandas as pd
        dt = pd.to_datetime(d, errors="coerce", dayfirst=True)
        if pd.isna(dt):
            return None
        return dt.to_pydatetime()
    except Exception:
        return None


def _get_month_name(month_key):
    """Convert YYYY-MM to 'Month Year' format."""
    try:
        dt = datetime.strptime(month_key, "%Y-%m")
        return dt.strftime("%B %Y")
    except:
        return month_key


def analyze_notes_for_progress(notes: List[Dict]) -> Dict[str, Any]:
    """Analyze notes for progress events and build timeline data."""
    results = {
        "total_notes": len(notes),
        "monthly_data": defaultdict(lambda: {
            "notes": [],
            "scores": [],
            "risk_factors": defaultdict(int),
            "tentpole_events": [],
            "violence_count": 0,
            "verbal_count": 0,
            "total_incidents": 0,
            "incidents": [],  # Store individual incidents for popup
        }),
        "tentpole_events": [],
        "all_months": set(),
        "monthly_violence": defaultdict(int),
        "monthly_verbal": defaultdict(int),
        "monthly_incidents": defaultdict(int),
        "episodes": [],  # Inpatient/community episodes from timeline builder
        "admissions": [],  # List of admission dates
        "discharges": [],  # List of discharge dates
    }

    # Build timeline to get admission/discharge episodes
    episodes = build_timeline(notes)
    results["episodes"] = episodes

    # Extract admission and discharge dates
    for ep in episodes:
        if ep.get("type") == "inpatient":
            results["admissions"].append({
                "date": ep["start"],
                "label": ep.get("label", "Admission"),
                "end": ep["end"],
            })
            # Discharge is the end of inpatient episode
            results["discharges"].append({
                "date": ep["end"],
                "label": ep.get("label", "Discharge"),
            })

    for note in notes:
        text = note.get("text", "") or note.get("content", "") or note.get("body", "")
        date = _normalise_date(note.get("date") or note.get("datetime"))

        if not text or not date:
            continue

        month_key = date.strftime("%Y-%m")
        results["all_months"].add(month_key)
        text_lower = text.lower()

        # Calculate risk score for this note
        score = 0
        for risk_name, risk_config in RISK_PATTERNS.items():
            for pattern in risk_config["patterns"]:
                match = re.search(pattern, text_lower)
                if match:
                    # Check for negative context (nil, no, denied, etc.)
                    if _has_negative_context(text_lower, match.start(), match.end()):
                        continue  # Skip this match - it's negated

                    score += risk_config["weight"]
                    results["monthly_data"][month_key]["risk_factors"][risk_name] += 1
                    results["monthly_data"][month_key]["total_incidents"] += 1
                    results["monthly_incidents"][month_key] += 1

                    # Store incident details for popup
                    results["monthly_data"][month_key]["incidents"].append({
                        "date": date,
                        "category": risk_name,
                        "matched": match.group(0),
                        "text": text[:200],
                        "severity": "high" if risk_config["weight"] >= 8 else "medium" if risk_config["weight"] >= 3 else "low",
                    })

                    # Track violence and verbal separately
                    if risk_name == "Physical Violence":
                        results["monthly_data"][month_key]["violence_count"] += 1
                        results["monthly_violence"][month_key] += 1
                    elif risk_name == "Verbal Aggression":
                        results["monthly_data"][month_key]["verbal_count"] += 1
                        results["monthly_verbal"][month_key] += 1
                    break

        results["monthly_data"][month_key]["scores"].append(score)
        results["monthly_data"][month_key]["notes"].append({"date": date, "text": text})

        # Check for tentpole events
        for event_name, event_config in TENTPOLE_PATTERNS.items():
            for pattern in event_config["patterns"]:
                if re.search(pattern, text_lower):
                    event = {
                        "date": date,
                        "month": month_key,
                        "type": event_name,
                        "color": event_config["color"],
                        "text": text,
                    }
                    results["tentpole_events"].append(event)
                    results["monthly_data"][month_key]["tentpole_events"].append(event)
                    break

    results["all_months"] = sorted(results["all_months"])
    return results


def identify_periods(results: Dict) -> List[Dict]:
    """Identify periods of stability, escalation, and high risk.

    Periods are split if they exceed 12 months to ensure manageable narrative chunks.
    """
    periods = []
    all_months = results["all_months"]
    monthly_data = results["monthly_data"]

    if not all_months:
        return periods

    current_period = None
    MAX_MONTHS = 12  # Never have a period longer than 12 months

    for month in all_months:
        data = monthly_data[month]
        scores = data["scores"]
        avg_score = sum(scores) / len(scores) if scores else 0

        # Classify month
        if avg_score >= 5:
            category = "high_risk"
        elif avg_score >= 2:
            category = "moderate"
        else:
            category = "stable"

        # Check if we need to start a new period
        start_new = False
        if current_period is None:
            start_new = True
        elif current_period["category"] != category:
            start_new = True
        elif len(current_period["months"]) >= MAX_MONTHS:
            # Split period if it exceeds max length
            start_new = True

        if start_new:
            if current_period:
                periods.append(current_period)
            current_period = {
                "category": category,
                "start": month,
                "end": month,
                "months": [month],
                "avg_score": avg_score,
            }
        else:
            current_period["end"] = month
            current_period["months"].append(month)
            # Update average
            all_scores = []
            for m in current_period["months"]:
                all_scores.extend(monthly_data[m]["scores"])
            current_period["avg_score"] = sum(all_scores) / len(all_scores) if all_scores else 0

    if current_period:
        periods.append(current_period)

    return periods


def generate_narrative(results: Dict, patient_name: str = "The patient", gender: str = None) -> str:
    """Generate a flowing narrative summary of the patient's journey.

    Creates one continuous narrative that flows through the timeline, mentioning
    incidents, admissions, and discharges naturally as they occur.

    Args:
        results: Analysis results from analyze_notes_for_progress
        patient_name: Patient's name
        gender: 'M', 'F', or None for neutral pronouns
    """
    from collections import defaultdict

    # Reset trackers for fresh narrative
    reset_phrase_tracker()
    reset_reference_tracker()

    all_months = results["all_months"]
    monthly_data = results["monthly_data"]
    tentpole_events = results["tentpole_events"]
    admissions = results.get("admissions", [])
    discharges = results.get("discharges", [])
    episodes = results.get("episodes", [])

    if not all_months:
        return "No data available for analysis."

    name = patient_name.split()[0] if patient_name else "The patient"

    # Set up pronouns based on gender
    if gender == 'F':
        pronoun = "she"
        pronoun_obj = "her"
        pronoun_poss = "her"
        pronoun_cap = "She"
        pronoun_poss_cap = "Her"
    elif gender == 'M':
        pronoun = "he"
        pronoun_obj = "him"
        pronoun_poss = "his"
        pronoun_cap = "He"
        pronoun_poss_cap = "His"
    else:
        pronoun = "they"
        pronoun_obj = "them"
        pronoun_poss = "their"
        pronoun_cap = "They"
        pronoun_poss_cap = "Their"

    # Date range
    start_date = _get_month_name(all_months[0])
    end_date = _get_month_name(all_months[-1])

    # Calculate overall statistics
    total_violence = sum(results.get("monthly_violence", {}).values())
    total_verbal = sum(results.get("monthly_verbal", {}).values())
    total_incidents = sum(results.get("monthly_incidents", {}).values())

    # Count inpatient vs community episodes
    inpatient_episodes = [ep for ep in episodes if ep.get("type") == "inpatient"]

    # Calculate total inpatient days
    total_inpatient_days = 0
    for ep in inpatient_episodes:
        start = ep.get("start")
        end = ep.get("end")
        if start and end:
            if hasattr(start, 'date'):
                start = start.date()
            if hasattr(end, 'date'):
                end = end.date()
            days = (end - start).days + 1
            total_inpatient_days += days

    narrative = f"""PROGRESS AND RISK NARRATIVE
{'='*60}

{patient_name}
{start_date} to {end_date}

OVERVIEW:
- Total months reviewed: {len(all_months)}
- Admissions: {len(inpatient_episodes)} ({total_inpatient_days} inpatient days total)
- Total incidents recorded: {total_incidents}
- Physical violence incidents: {total_violence}
- Verbal aggression incidents: {total_verbal}
- Key events documented: {len(tentpole_events)}

"""

    # Add admission/discharge summary if any
    if inpatient_episodes:
        narrative += "ADMISSION HISTORY:\n"
        for i, ep in enumerate(inpatient_episodes, 1):
            start = ep.get("start")
            end = ep.get("end")
            if start and end:
                if hasattr(start, 'date'):
                    start = start.date()
                if hasattr(end, 'date'):
                    end = end.date()
                days = (end - start).days + 1
                start_str = start.strftime("%d %b %Y") if hasattr(start, 'strftime') else str(start)
                end_str = end.strftime("%d %b %Y") if hasattr(end, 'strftime') else str(end)
                narrative += f"- Admission {i}: {start_str} to {end_str} ({days} days)\n"
        narrative += "\n"

    narrative += "NARRATIVE:\n\n"

    # Collect all incidents with dates
    all_incidents_list = []
    for month in all_months:
        data = monthly_data.get(month, {})
        for inc in data.get("incidents", []):
            if inc.get("date"):
                all_incidents_list.append(inc)

    # Sort incidents by date
    all_incidents_list.sort(key=lambda x: x["date"])

    # Create admission/discharge events with dates
    admission_events = []
    for adm in admissions:
        adm_date = adm.get("date")
        if adm_date:
            if hasattr(adm_date, 'date'):
                adm_date = adm_date.date()
            admission_events.append({
                "date": adm_date,
                "type": "admission",
                "end": adm.get("end"),
            })

    discharge_events = []
    for dis in discharges:
        dis_date = dis.get("date")
        if dis_date:
            if hasattr(dis_date, 'date'):
                dis_date = dis_date.date()
            discharge_events.append({
                "date": dis_date,
                "type": "discharge",
            })

    # Helper to get year from date
    def get_year(d):
        if hasattr(d, 'year'):
            return d.year
        return None

    # Helper to format date
    def fmt_date(d):
        if hasattr(d, 'strftime'):
            return d.strftime("%d %B %Y")
        return str(d)

    def fmt_date_short(d):
        if hasattr(d, 'strftime'):
            return d.strftime("%d %b")
        return str(d)

    # Group incidents by year
    incidents_by_year = defaultdict(list)
    for inc in all_incidents_list:
        year = get_year(inc["date"])
        if year:
            incidents_by_year[year].append(inc)

    # Get the range of years
    years = sorted(set(get_year(datetime.strptime(m, "%Y-%m")) for m in all_months))

    # Check if a date is during inpatient
    def is_inpatient(check_date):
        if not check_date:
            return False
        check_d = check_date.date() if hasattr(check_date, 'date') else check_date
        for ep in episodes:
            if ep.get("type") == "inpatient":
                start = ep.get("start")
                end = ep.get("end")
                if start and end:
                    start_d = start.date() if hasattr(start, 'date') else start
                    end_d = end.date() if hasattr(end, 'date') else end
                    if start_d <= check_d <= end_d:
                        return True
        return False

    # Find admissions in a year
    def get_admissions_in_year(year):
        return [a for a in admission_events if get_year(a["date"]) == year]

    def get_discharges_in_year(year):
        return [d for d in discharge_events if get_year(d["date"]) == year]

    def is_year_covered_by_ongoing_admission(year):
        """Check if a year falls within an admission that started in a previous year
        AND there are no NEW admissions starting in this year.
        This helps avoid generating separate yearly summaries for years that are
        already covered by a long inpatient stay, while still processing new admissions."""
        # First check if there are new admissions starting in this year
        year_admits = get_admissions_in_year(year)
        if year_admits:
            return False  # Always process years with new admissions

        for ep in episodes:
            if ep.get("type") == "inpatient":
                start = ep.get("start")
                end = ep.get("end")
                if start and end:
                    start_year = get_year(start)
                    end_year = get_year(end)
                    # Year is covered if admission started before this year and ends during/after
                    if start_year < year <= end_year:
                        return True
        return False

    # Pre-process: classify years and group consecutive similar years
    def classify_year(y):
        """Classify a year as 'stable', 'minor', or 'significant'."""
        incidents = incidents_by_year.get(y, [])
        admissions = get_admissions_in_year(y)
        if admissions:
            return 'admission'
        elif len(incidents) == 0:
            return 'stable'
        elif len(incidents) <= 3:
            return 'minor'
        else:
            return 'significant'

    # Group consecutive years with same classification (stable or minor)
    # This avoids repetitive "remained stable through X" phrases
    year_groups = []
    i = 0
    while i < len(years):
        year_class = classify_year(years[i])

        if year_class in ('stable', 'minor'):
            # Start a group of stable/minor years
            group_start = years[i]
            group_end = years[i]
            group_incidents = list(incidents_by_year.get(years[i], []))

            # Extend group while next year is also stable or minor
            while i + 1 < len(years):
                next_class = classify_year(years[i + 1])
                if next_class in ('stable', 'minor'):
                    i += 1
                    group_end = years[i]
                    group_incidents.extend(incidents_by_year.get(years[i], []))
                else:
                    break

            year_groups.append({
                'type': 'stable_minor',
                'start': group_start,
                'end': group_end,
                'incidents': group_incidents,
                'years': list(range(group_start, group_end + 1)),
            })
        else:
            # Admission or significant year - keep separate
            year_groups.append({
                'type': year_class,
                'start': years[i],
                'end': years[i],
                'incidents': list(incidents_by_year.get(years[i], [])),
                'years': [years[i]],
            })
        i += 1

    # Start the flowing narrative
    first_year = years[0] if years else None
    if first_year:
        # Determine if starting in community or inpatient
        first_month = all_months[0]
        first_date = datetime.strptime(first_month, "%Y-%m")

        if is_inpatient(first_date):
            narrative += f"At the start of the records in {start_date}, {name} was an inpatient. "
        else:
            narrative += f"The records begin in {start_date} when {name} was in the community"
            first_year_incidents = incidents_by_year.get(first_year, [])
            if not first_year_incidents:
                narrative += f" and {pronoun} presented as stable with no significant concerns documented. "
            else:
                narrative += ". "

    # Process groups, creating flowing prose
    last_mentioned_state = "stable"
    processed_admissions = set()  # Track admission dates we've already described
    admission_number = 0  # Track admission count for numbering
    last_discharge_date = None  # Track previous discharge date for gap calculation

    # Ordinal words for admission numbering
    ORDINALS = [
        "first", "second", "third", "fourth", "fifth", "sixth", "seventh",
        "eighth", "ninth", "tenth", "eleventh", "twelfth", "thirteenth",
        "fourteenth", "fifteenth", "sixteenth", "seventeenth", "eighteenth",
        "nineteenth", "twentieth"
    ]

    def get_ordinal(n):
        """Get ordinal word for number (1-indexed)."""
        if 1 <= n <= len(ORDINALS):
            return ORDINALS[n - 1]
        else:
            # Fallback to numeric ordinal
            if n % 10 == 1 and n != 11:
                return f"{n}st"
            elif n % 10 == 2 and n != 12:
                return f"{n}nd"
            elif n % 10 == 3 and n != 13:
                return f"{n}rd"
            else:
                return f"{n}th"

    for group_idx, group in enumerate(year_groups):
        group_type = group['type']
        group_start = group['start']
        group_end = group['end']
        group_incidents = group['incidents']

        # Skip the first year if already mentioned in opening
        if group_idx == 0 and first_year and group_start == first_year:
            if group_type == 'stable_minor' and not group_incidents:
                continue  # Already said "presented as stable"

        if group_type == 'stable_minor':
            # Handle grouped stable/minor years
            if not group_incidents:
                # Completely stable period - use varied phrasing
                if group_start == group_end:
                    stable_phrase = get_phrase("remained_stable", pronoun=pronoun, pronoun_poss=pronoun_poss)
                    narrative += f"{stable_phrase} throughout {group_start}. "
                else:
                    year_desc = f"{group_start}-{group_end}"
                    stable_phrase = get_phrase("remained_stable", pronoun=pronoun, pronoun_poss=pronoun_poss)
                    narrative += f"{stable_phrase} throughout {year_desc}. "
                last_mentioned_state = "stable"
            else:
                # Minor incidents across this period
                # Collect all incidents with dates
                all_incident_descs = []
                for inc in sorted(group_incidents, key=lambda x: x.get('date') or datetime.min):
                    inc_date = inc.get('date')
                    inc_cat = inc.get('category', '')
                    inc_matched = inc.get('matched', '')
                    if inc_date:
                        date_str = fmt_date_short(inc_date) + f" {inc_date.year}"
                        if inc_cat == "Physical Violence":
                            link = make_link("physical aggression", inc_date, inc_matched or "physical aggression")
                            all_incident_descs.append(f"{link} ({date_str})")
                        elif inc_cat == "Verbal Aggression":
                            link = make_link("verbal aggression", inc_date, inc_matched or "verbal aggression")
                            all_incident_descs.append(f"{link} ({date_str})")
                        elif inc_cat == "Self-Harm":
                            link = make_link("self-harm", inc_date, inc_matched or "self-harm")
                            all_incident_descs.append(f"{link} ({date_str})")

                if group_start == group_end:
                    year_desc = str(group_start)
                else:
                    year_desc = f"{group_start}-{group_end}"

                if len(all_incident_descs) <= 4:
                    incidents_text = ", ".join(all_incident_descs)
                else:
                    incidents_text = ", ".join(all_incident_descs[:4]) + f" and {len(all_incident_descs) - 4} further concerns"

                stable_phrase = get_phrase("remained_stable", pronoun=pronoun, pronoun_poss=pronoun_poss)
                narrative += f"{stable_phrase} through {year_desc} with isolated concerns: {incidents_text}. "
                last_mentioned_state = "stable"
            continue

        # For admission years, get the year data
        year = group_start
        year_incidents = incidents_by_year.get(year, [])
        year_admissions = get_admissions_in_year(year)
        year_discharges = get_discharges_in_year(year)

        # Group incidents by type
        violence_incidents = [i for i in year_incidents if i.get("category") == "Physical Violence"]
        verbal_incidents = [i for i in year_incidents if i.get("category") == "Verbal Aggression"]
        self_harm_incidents = [i for i in year_incidents if i.get("category") == "Self-Harm"]

        total_year_incidents = len(year_incidents)

        # Skip years that are covered by an ongoing admission from a previous year
        # These incidents are already described in the inpatient stay narrative
        if is_year_covered_by_ongoing_admission(year):
            continue  # Skip this year entirely - it's part of a multi-year admission

        # Describe significant years or admission years
        if group_type == 'significant':
            if last_mentioned_state == "stable":
                narrative += get_phrase("year_deterioration", year=year, pronoun=pronoun, pronoun_poss=pronoun_poss)
            else:
                narrative += get_phrase("year_continued_instability", year=year)

        # Mention incidents for this year (if not an admission year - those get handled separately)
        # Use comparative language rather than numbers
        if not year_admissions:
            # Get previous year's incidents for comparison
            prev_year = year - 1
            prev_year_incidents = incidents_by_year.get(prev_year, [])
            prev_violence = len([i for i in prev_year_incidents if i.get("category") == "Physical Violence"])
            prev_verbal = len([i for i in prev_year_incidents if i.get("category") == "Verbal Aggression"])
            prev_self_harm = len([i for i in prev_year_incidents if i.get("category") == "Self-Harm"])

            incident_parts = []

            if violence_incidents:
                vi = violence_incidents[0]  # Use first incident for link
                vi_date = vi['date']
                vi_matched = vi.get('matched', 'physical aggression')
                vi_snippet = vi.get('text', '')[:100]
                if len(violence_incidents) == 1:
                    link = make_link("physical aggression", vi_date, vi_matched, vi_snippet)
                    incident_parts.append(f"{link} on {fmt_date(vi_date)}")
                elif prev_violence > 0:
                    # Compare to previous year
                    if len(violence_incidents) > prev_violence * 1.3:
                        incident_parts.append(make_link("increased physical aggression", vi_date, vi_matched, vi_snippet))
                    elif len(violence_incidents) < prev_violence * 0.7:
                        incident_parts.append(make_link("reduced physical aggression", vi_date, vi_matched, vi_snippet))
                    else:
                        incident_parts.append(make_link("ongoing physical aggression", vi_date, vi_matched, vi_snippet))
                else:
                    incident_parts.append(make_link("physical aggression", vi_date, vi_matched, vi_snippet))

            if verbal_incidents:
                vb = verbal_incidents[0]  # Use first incident for link
                vb_date = vb['date']
                vb_matched = vb.get('matched', 'verbal aggression')
                vb_snippet = vb.get('text', '')[:100]
                if len(verbal_incidents) == 1:
                    link = make_link("verbal aggression", vb_date, vb_matched, vb_snippet)
                    incident_parts.append(f"{link} on {fmt_date_short(vb_date)}")
                elif prev_verbal > 0:
                    if len(verbal_incidents) > prev_verbal * 1.3:
                        incident_parts.append(make_link("increased verbal aggression", vb_date, vb_matched, vb_snippet))
                    elif len(verbal_incidents) < prev_verbal * 0.7:
                        incident_parts.append(make_link("reduced verbal aggression", vb_date, vb_matched, vb_snippet))
                    else:
                        incident_parts.append(make_link("ongoing verbal aggression", vb_date, vb_matched, vb_snippet))
                else:
                    incident_parts.append(make_link("verbal aggression", vb_date, vb_matched, vb_snippet))

            if self_harm_incidents:
                sh = self_harm_incidents[0]  # Use first incident for link
                sh_date = sh['date']
                sh_matched = sh.get('matched', 'self-harm')
                sh_snippet = sh.get('text', '')[:100]
                if len(self_harm_incidents) == 1:
                    link = make_link("self-harm", sh_date, sh_matched, sh_snippet)
                    incident_parts.append(f"{link} on {fmt_date(sh_date)}")
                elif len(self_harm_incidents) <= 3:
                    sh_dates = [fmt_date_short(i["date"]) for i in self_harm_incidents[:3]]
                    link = make_link("self-harm", sh_date, sh_matched, sh_snippet)
                    incident_parts.append(f"{link} on {', '.join(sh_dates)}")
                elif prev_self_harm > 0:
                    if len(self_harm_incidents) > prev_self_harm * 1.3:
                        incident_parts.append(make_link("increased self-harm", sh_date, sh_matched, sh_snippet))
                    elif len(self_harm_incidents) < prev_self_harm * 0.7:
                        incident_parts.append(make_link("reduced self-harm", sh_date, sh_matched, sh_snippet))
                    else:
                        incident_parts.append(make_link("ongoing self-harm", sh_date, sh_matched, sh_snippet))
                else:
                    incident_parts.append(make_link("self-harm", sh_date, sh_matched, sh_snippet))

            if incident_parts and group_idx > 0:
                if total_year_incidents <= 2:
                    narrative += f" with only {' and '.join(incident_parts)}. "
                else:
                    narrative += f" with {' and '.join(incident_parts)}. "
                last_mentioned_state = "incidents"
            elif group_idx > 0 and group_type == 'significant':
                narrative += ". "
                last_mentioned_state = "incidents"

        # Handle admissions and discharges - each admission starts a new paragraph
        # Sort all admission events to find next admission after each discharge
        all_admissions_sorted = sorted(admission_events, key=lambda x: x["date"])

        # Get all notes for detailed extraction (from shared store)
        all_notes = get_shared_store().notes

        for adm_idx, adm in enumerate(year_admissions):
            adm_date = adm["date"]

            # Skip if we've already processed this admission
            adm_key = str(adm_date)
            if adm_key in processed_admissions:
                continue
            processed_admissions.add(adm_key)

            adm_end = adm.get("end")
            if adm_end and hasattr(adm_end, 'date'):
                adm_end = adm_end.date()

            # Start new paragraph for admission
            narrative += "\n\n"

            # Increment admission number
            admission_number += 1
            ordinal = get_ordinal(admission_number)
            admission_label = f"**{ordinal} admission**"

            # Extract detailed admission information
            adm_details = extract_admission_details(all_notes, adm_date, adm_end, episodes)

            # Build admission description with varied phrasing
            if adm_end:
                days = (adm_end - adm_date).days + 1
                if days < 30:
                    duration_desc = "less than a month"
                elif days < 45:
                    duration_desc = f"about {days} days"
                elif days < 75:
                    weeks = round(days / 7)
                    duration_desc = f"approximately {weeks} weeks"
                else:
                    months = round(days / 30)
                    duration_desc = f"approximately {months} months"

                # Describe admission with rich context from first 2 weeks of notes
                duration_intro = get_phrase("duration_intro")
                ctx = adm_details.get('admission_context', {})

                # Build admission intro
                if admission_number == 1:
                    narrative += f"The {admission_label} was on {fmt_date(adm_date)}. "
                else:
                    # Calculate gap from previous discharge if available
                    if last_discharge_date and adm_date:
                        gap_days = (adm_date - last_discharge_date).days
                        if gap_days < 14:
                            gap_desc = f"{gap_days} days"
                        elif gap_days < 60:
                            weeks = round(gap_days / 7)
                            gap_desc = f"{weeks} week{'s' if weeks != 1 else ''}"
                        elif gap_days < 365:
                            months = round(gap_days / 30)
                            gap_desc = f"{months} month{'s' if months != 1 else ''}"
                        else:
                            years_approx = round(gap_days / 365, 1)
                            if years_approx == int(years_approx):
                                years_approx = int(years_approx)
                            gap_desc = f"approximately {years_approx} year{'s' if years_approx != 1 else ''}"

                        # Combine gap description with admission statement
                        if gap_days <= 60:
                            narrative += f"{pronoun_cap} had been out of hospital for only {gap_desc} before {pronoun_poss} {admission_label} on {fmt_date(adm_date)}. "
                        else:
                            narrative += f"{pronoun_cap} had been in the community for {gap_desc} before {pronoun_poss} {admission_label} on {fmt_date(adm_date)}. "
                        # Skip the separate admission date statement since it's now combined
                        gap_combined_with_admission = True
                    else:
                        gap_combined_with_admission = False

                    # Check if triggers repeat - only add separate admission statement if not combined above
                    if not gap_combined_with_admission:
                        triggers = adm_details['triggers'][:2] if adm_details['triggers'] else []
                        has_repeat = any(_phrase_tracker.is_repeated_trigger(t) for t in triggers)
                        if has_repeat:
                            again_phrase = get_again_phrase("admission_again", pronoun=pronoun, pronoun_cap=pronoun_cap)
                            narrative += f"{again_phrase} was admitted for the {admission_label} on {fmt_date(adm_date)}. "
                        else:
                            narrative += f"The {admission_label} commenced on {fmt_date(adm_date)}. "

                # Source of admission (A&E, CTO recall, etc.)
                if ctx.get('source'):
                    if ctx['source'] == 'A&E':
                        narrative += f"{pronoun_cap} had presented to A&E "
                    elif ctx['source'] == 'CTO recall':
                        narrative += f"{pronoun_cap} was recalled under {pronoun_poss} CTO "
                    elif ctx['source'] == 'Section 136 (police)':
                        narrative += f"{pronoun_cap} was brought in under Section 136 "
                    elif ctx['source'] == 'the community':
                        # Vary the phrasing for community admissions
                        community_phrases = [
                            f"{pronoun_cap} was admitted from the community",
                            f"{pronoun_cap} presented from the community",
                            f"This followed deterioration in the community",
                            f"{pronoun_cap} required admission from community care",
                        ]
                        narrative += f"{random.choice(community_phrases)} "
                    else:
                        narrative += f"{pronoun_cap} was referred from {ctx['source']} "

                    # Add presenting complaints - always list symptoms, vary phrasing if repeated
                    if ctx.get('presenting_complaint'):
                        raw_complaints = ctx['presenting_complaint'][:3]
                        # Extract labels for phrase tracker (handle both tuple and string formats)
                        complaint_labels = [c[0] if isinstance(c, tuple) else c for c in raw_complaints]
                        # Wrap each complaint in a clickable link using actual matched text, date, and content snippet
                        linked_complaints = []
                        for c in raw_complaints:
                            if isinstance(c, tuple) and len(c) >= 4:
                                # New format: (label, matched, note_date, content_snippet)
                                label, matched, note_date, content_snippet = c
                                linked_complaints.append(make_link(label, note_date, matched, content_snippet))
                            elif isinstance(c, tuple) and len(c) >= 3:
                                # Old format: (label, matched, note_date)
                                label, matched, note_date = c
                                linked_complaints.append(make_link(label, note_date, matched))
                            else:
                                # Fallback for string format
                                label = c[0] if isinstance(c, tuple) else c
                                linked_complaints.append(make_link(label, adm_date, label))
                        symptoms_text = ', '.join(linked_complaints)
                        if _phrase_tracker.are_complaints_similar(complaint_labels):
                            # Use varied phrase with symptoms always listed
                            phrase = get_phrase("repeated_presentation",
                                                pronoun=pronoun,
                                                pronoun_cap=pronoun_cap,
                                                pronoun_poss=pronoun_poss,
                                                symptoms=symptoms_text)
                            narrative += f"{phrase}. "
                        else:
                            narrative += f"with {symptoms_text}. "
                        _phrase_tracker.record_complaints(complaint_labels)
                    else:
                        narrative += ". "
                elif adm_details['triggers']:
                    # Fall back to triggers if no rich context
                    triggers = adm_details['triggers'][:2]
                    # Wrap each trigger in a clickable link (triggers are tuples with dates)
                    linked_triggers = [make_link(t[0] if isinstance(t, tuple) else t,
                                                 t[2] if isinstance(t, tuple) and len(t) > 2 else adm_date,
                                                 t[0] if isinstance(t, tuple) else t)
                                       for t in triggers]
                    trigger_text = ", ".join(linked_triggers)
                    narrative += f"{pronoun_cap} presented with {trigger_text}. "

                    # Record triggers for repeat detection
                    for t in triggers:
                        _phrase_tracker.record_trigger(t)

                # Legal status - with clickable reference
                if ctx.get('legal_status'):
                    legal = ctx['legal_status']
                    # Handle both old string format and new dict format
                    if isinstance(legal, dict):
                        status = legal.get('status', '')
                        legal_link = make_link(status, legal.get('date'), legal.get('matched', status), legal.get('snippet', ''))
                        if 'Section 3' in status:
                            narrative += f"{pronoun_cap} was detained under {legal_link} of the Mental Health Act. "
                        elif 'Section 2' in status:
                            narrative += f"{pronoun_cap} was detained under {legal_link} for assessment. "
                        elif 'CTO revoked' in status:
                            narrative += f"{pronoun_poss_cap} CTO was revoked and {pronoun} was detained under {legal_link}. "
                    else:
                        # Fallback for old string format
                        if 'Section 3' in legal:
                            narrative += f"{pronoun_cap} was detained under Section 3 of the Mental Health Act. "
                        elif 'Section 2' in legal:
                            narrative += f"{pronoun_cap} was detained under Section 2 for assessment. "
                        elif 'CTO revoked' in legal:
                            narrative += f"{pronoun_poss_cap} CTO was revoked and {pronoun} was detained under Section 3. "

                # Medication status on admission
                if ctx.get('medication_status'):
                    med_status = ctx['medication_status']
                    if isinstance(med_status, dict):
                        status_text = med_status.get('text', '')
                        status_date = med_status.get('date')
                        status_link = make_link(status_text, status_date, med_status.get('matched', status_text), med_status.get('snippet', ''))

                        # Check if the evidence is actually near admission (within 5 days)
                        # If not, don't say "at the time of admission"
                        is_near_admission = False
                        if status_date and adm_date:
                            days_from_admission = (status_date - adm_date).days
                            is_near_admission = days_from_admission <= 5

                        if 'off Clozapine' in status_text:
                            if is_near_admission:
                                narrative += f"At the time of admission, {pronoun} had been {status_link}. "
                            else:
                                narrative += f"{pronoun_cap} had been {status_link}. "
                        elif 'refusing' in status_text:
                            if is_near_admission:
                                narrative += f"{pronoun_cap} was {status_link} at the time of admission. "
                            else:
                                # Use actual date if not near admission
                                date_str = fmt_date(status_date) if status_date else "during the admission"
                                narrative += f"On {date_str}, {pronoun} was noted to be {status_link}. "
                        else:
                            narrative += f"{pronoun_cap} was {status_link}. "
                    else:
                        # Old string format fallback
                        if 'off Clozapine' in med_status:
                            narrative += f"At the time of admission, {pronoun} had been {med_status}. "
                        elif 'refusing' in med_status:
                            narrative += f"{pronoun_cap} was {med_status} at the time of admission. "
                        else:
                            narrative += f"{pronoun_cap} was {med_status}. "
                elif adm_details['medications_before']:
                    meds_text = format_meds_as_links(adm_details['medications_before'], limit=4)
                    narrative += f"On admission, {pronoun} was on {meds_text}. "

                # Key concerns with repetition checking
                if ctx.get('key_concerns'):
                    concerns = ctx['key_concerns'][:2]
                    # Extract text for repetition checking
                    concern_texts = [c['text'] if isinstance(c, dict) else c for c in concerns]
                    if _phrase_tracker.are_concerns_repeated(concern_texts):
                        # Vary phrasing for repeated concerns
                        repeated_concern_phrases = [
                            "The same concerns applied regarding her care.",
                            "Clinical concerns remained unchanged.",
                            "The familiar concerns about her treatment capacity persisted.",
                            "",  # Sometimes skip entirely if repetitive
                        ]
                        phrase = random.choice(repeated_concern_phrases)
                        if phrase:
                            narrative += f"{phrase} "
                    else:
                        concern_links = []
                        for c in concerns:
                            if isinstance(c, dict):
                                link = make_link(c['text'], c.get('date'), c.get('matched', c['text']), c.get('content_snippet', ''))
                                concern_links.append(link)
                            else:
                                concern_links.append(c)
                        narrative += f"Key concerns included {' and '.join(concern_links)}. "
                    _phrase_tracker.record_concerns(concern_texts)

                # Substance misuse details if any
                if ctx.get('substance_misuse'):
                    subs = ctx['substance_misuse']
                    sub_texts = []
                    for s in subs[:3]:  # Max 3 substances
                        sub_date = s.get('date')
                        sub_snippet = s.get('content_snippet', '')
                        sub_matched = s.get('matched', s['substance'])
                        if s.get('frequency') and s.get('amount'):
                            sub_text = f"{s['frequency']} {s['substance']} ({s['amount']})"
                        elif s.get('frequency'):
                            sub_text = f"{s['frequency']} {s['substance']} use"
                        elif s.get('amount'):
                            sub_text = f"{s['substance']} ({s['amount']})"
                        else:
                            sub_text = f"{s['substance']} use"
                        sub_link = make_link(sub_text, sub_date, sub_matched, sub_snippet)
                        sub_texts.append(sub_link)
                    if sub_texts:
                        narrative += f"There was documented {', '.join(sub_texts)}. "

                # Duration
                narrative += f"This admission lasted {duration_desc}. "

                # Describe events during admission with progression
                if adm_details['incidents']:
                    during_phrase = get_phrase("during_admission")
                    incident_count = len(adm_details['incidents'])
                    progression = adm_details.get('incident_progression', {})
                    pattern = progression.get('pattern', 'unknown')

                    # Get incident types ordered by frequency (without counts)
                    top_incidents = sorted(adm_details['incident_summary'].items(), key=lambda x: -x[1])
                    main_concern_types = [t for t, c in top_incidents[:3]]

                    # Convert to links by finding first incident of each type
                    main_concerns = []
                    for concern_type in main_concern_types:
                        # Find first incident of this type
                        first_inc = next((i for i in adm_details['incidents'] if i.get('type') == concern_type), None)
                        if first_inc:
                            concern_link = make_link(concern_type, first_inc.get('date'), concern_type, first_inc.get('summary', '')[:100])
                            main_concerns.append(concern_link)
                        else:
                            main_concerns.append(concern_type)

                    # Qualitative description based on count
                    if incident_count <= 3:
                        quantity_desc = "a small number of concerns"
                    elif incident_count <= 10:
                        quantity_desc = "some concerns"
                    elif incident_count <= 50:
                        quantity_desc = "multiple concerns"
                    elif incident_count <= 200:
                        quantity_desc = "many concerns"
                    else:
                        quantity_desc = "frequent concerns"

                    # Describe progression pattern qualitatively (duration-aware language)
                    if pattern == 'improving' and incident_count > 5:
                        if days > 60:
                            narrative += f"{during_phrase}, the first two months were marked by {quantity_desc}, "
                        elif days > 21:
                            narrative += f"{during_phrase}, the initial weeks were marked by {quantity_desc}, "
                        elif days > 7:
                            narrative += f"{during_phrase}, initially there were {quantity_desc}, "
                        else:
                            # Very short admission - just describe incidents without "initially"
                            narrative += f"{during_phrase}, there were {quantity_desc}, "
                        if main_concerns:
                            narrative += f"mainly concerns about {' and '.join(main_concerns)}. "
                        else:
                            narrative += ". "
                        if days > 7:
                            narrative += f"These gradually reduced in the latter part of the admission. "

                    elif pattern == 'worsening' and incident_count > 5:
                        if days > 14:
                            narrative += f"{during_phrase}, {pronoun_poss} behaviour initially appeared to settle but subsequently deteriorated "
                        else:
                            narrative += f"{during_phrase}, {pronoun_poss} behaviour deteriorated "
                        if main_concerns:
                            narrative += f"with concerns about {' and '.join(main_concerns)}. "
                        else:
                            narrative += ". "

                    elif pattern == 'fluctuating' and incident_count > 5:
                        narrative += f"{during_phrase}, {pronoun_poss} presentation fluctuated with {quantity_desc}, "
                        if main_concerns:
                            narrative += f"mainly concerns about {' and '.join(main_concerns)}. "
                        else:
                            narrative += ". "

                    elif incident_count <= 5:
                        narrative += f"{during_phrase}, there were {quantity_desc}. "
                        if main_concerns:
                            narrative += f"These involved {', '.join(main_concerns)}. "

                    else:
                        # Default description
                        narrative += f"{during_phrase}, there were {quantity_desc}, "
                        if main_concerns:
                            narrative += f"mainly {' and '.join(main_concerns)}. "
                        else:
                            narrative += ". "

                    # Add notable incidents (seclusion, 1-1, police, etc.) as flowing narrative
                    notable = adm_details.get('notable_incidents', [])
                    if notable:
                        # Sort by date
                        notable = sorted(notable, key=lambda x: x['date'])[:4]

                        # Map reason types to readable text
                        reason_text_map = {
                            'assault': 'an assault on staff',
                            'self-harm': 'self-harm',
                            'aggression': 'aggressive behaviour',
                            'agitation': 'significant agitation',
                            'absconding': 'an attempt to abscond',
                            'non-compliance': 'treatment refusal',
                        }

                        # Helper to format incident type for narrative
                        def format_incident_type(inc_type):
                            if inc_type == 'seclusion':
                                return "seclusion"
                            return inc_type

                        # Helper to create clickable incident type
                        def linked_incident_type(inc_type, inc_date, matched="", content_snippet=""):
                            formatted = format_incident_type(inc_type)
                            return make_link(formatted, inc_date, matched or formatted, content_snippet)

                        # Helper to clean up raw reason text
                        def clean_reason(reason_text, incident_type=None):
                            if not reason_text:
                                return ""
                            # Remove pronouns at start that don't fit grammatically
                            reason_text = re.sub(r'^(she|he|they|patient)\s+', '', reason_text, flags=re.IGNORECASE)
                            # Remove leading articles if followed by verb
                            reason_text = re.sub(r'^(the|a|an)\s+(?=\w+ed\b|\w+ing\b)', '', reason_text, flags=re.IGNORECASE)

                            # Reject fragments that don't make sense as standalone clauses
                            invalid_patterns = [
                                r'\bbut\s+(less|more|not|no)\b',  # "but less frequently"
                                r'\bor\s+(less|more|not|no)\b',   # "or more"
                                r'\bhowever\b',
                                r'\balthough\b',
                                r'\bwhereas\b',
                                r'\bdespite\b',
                                r'^(and|or|but)\s',  # Starting with conjunction
                                r'\btaking\s+prn\b',  # Medication actions
                                r'\bmedication\s+(given|administered)\b',
                                r'\brequested\s+(prn|medication)\b',
                                r'\bstaff\s+(spent|with|support)\b',  # Staff activities
                                r'\btime\s+with\b',  # "time with her"
                                r'\bseveral\s+occasions\b',  # Vague
                            ]
                            for pattern in invalid_patterns:
                                if re.search(pattern, reason_text, re.IGNORECASE):
                                    return ""

                            # SECLUSION requires violence/assault/agitation - strict validation
                            if incident_type == 'seclusion':
                                seclusion_valid_reasons = r'\b(assault|attack|hit|punch|kick|bit|bite|violen|aggress|threaten|threw|throw|spit|strangle|choke|weapon|self[- ]?harm|ligature|restrain|agitat|escalat|risk|disturb|destruct)\b'
                                if not re.search(seclusion_valid_reasons, reason_text, re.IGNORECASE):
                                    return ""  # Reject non-violence reasons for seclusion

                            # Must start with a noun, verb-ing, or clear action
                            if not re.match(r'^[A-Za-z]', reason_text):
                                return ""

                            return reason_text.strip()

                        # Build narrative description
                        if len(notable) == 1:
                            n = notable[0]
                            date_str = n['date'].strftime('%d %B %Y')
                            reason = n.get('reason', '')
                            desc = n.get('description', '')
                            inc_type = linked_incident_type(n['type'], n['date'], n.get('matched', ''), n.get('content_snippet', ''))

                            # Check if incident is on admission day
                            is_admission_day = (n['date'] == adm_date)

                            if is_admission_day:
                                # Same-day admission incident - special phrasing
                                if reason and reason in reason_text_map:
                                    narrative += f"{pronoun_cap} required {inc_type} on admission following {reason_text_map[reason]}. "
                                elif reason:
                                    clean_r = clean_reason(reason, n['type'])
                                    if clean_r:
                                        narrative += f"{pronoun_cap} required {inc_type} on admission following {clean_r}. "
                                    else:
                                        narrative += f"{pronoun_cap} required {inc_type} on admission. "
                                else:
                                    narrative += f"{pronoun_cap} required {inc_type} on admission. "
                            elif reason and reason in reason_text_map:
                                narrative += f"On {date_str}, {reason_text_map[reason]} resulted in {pronoun_obj} requiring {inc_type}. "
                            elif reason:
                                clean_r = clean_reason(reason, n['type'])
                                if clean_r:
                                    narrative += f"On {date_str}, {clean_r} resulted in {pronoun_obj} requiring {inc_type}. "
                                else:
                                    narrative += f"On {date_str}, there was a concern requiring {inc_type}. "
                            else:
                                narrative += f"On {date_str}, there was a concern requiring {inc_type}. "

                        else:
                            # Multiple incidents - check if they're consecutive (within 3 days)
                            first = notable[0]
                            first_date = first['date']
                            first_date_str = first_date.strftime('%d %B %Y')
                            first_reason = first.get('reason', '')
                            first_type = linked_incident_type(first['type'], first_date, first.get('matched', ''), first.get('content_snippet', ''))

                            # Check if first incident is on admission day
                            is_first_on_admission = (first_date == adm_date)

                            # Describe first incident
                            if is_first_on_admission:
                                # Same-day admission incident
                                if first_reason and first_reason in reason_text_map:
                                    narrative += f"{pronoun_cap} required {first_type} on admission following {reason_text_map[first_reason]}. "
                                elif first_reason:
                                    clean_r = clean_reason(first_reason, first['type'])
                                    if clean_r:
                                        narrative += f"{pronoun_cap} required {first_type} on admission following {clean_r}. "
                                    else:
                                        narrative += f"{pronoun_cap} required {first_type} on admission. "
                                else:
                                    narrative += f"{pronoun_cap} required {first_type} on admission. "
                            elif first_reason and first_reason in reason_text_map:
                                narrative += f"On {first_date_str}, {reason_text_map[first_reason]} led to {pronoun_obj} requiring {first_type}. "
                            elif first_reason:
                                clean_r = clean_reason(first_reason, first['type'])
                                if clean_r:
                                    narrative += f"On {first_date_str}, {clean_r} resulted in {pronoun_obj} requiring {first_type}. "
                                else:
                                    narrative += f"On {first_date_str}, there was a concern requiring {first_type}. "
                            else:
                                narrative += f"On {first_date_str}, there was a concern requiring {first_type}. "

                            # Check for escalation pattern
                            escalated_to_seclusion = any(n['type'] == 'seclusion' for n in notable[1:])
                            consecutive_days = all(
                                (notable[i+1]['date'] - notable[i]['date']).days <= 3
                                for i in range(len(notable)-1)
                            )

                            if consecutive_days and len(notable) > 1:
                                last = notable[-1]
                                last_date_str = last['date'].strftime('%d %B %Y')
                                last_type = linked_incident_type(last['type'], last['date'], last.get('matched', ''), last.get('content_snippet', ''))

                                if escalated_to_seclusion and first['type'] != 'seclusion':
                                    # Escalation narrative
                                    seclusion_link = linked_incident_type('seclusion', last['date'], last.get('matched', ''), last.get('content_snippet', ''))
                                    if len(notable) == 2:
                                        narrative += f"The situation escalated, resulting in {pronoun_obj} requiring {seclusion_link} on {last_date_str}. "
                                    else:
                                        mid_incidents = len(notable) - 2
                                        narrative += f"Difficulties continued over the following days with {mid_incidents} further concern{'s' if mid_incidents > 1 else ''}, "
                                        narrative += f"ultimately requiring {seclusion_link} on {last_date_str}. "
                                else:
                                    # Continuation narrative
                                    days_span = (last['date'] - first_date).days
                                    if days_span <= 1:
                                        narrative += f"Further concerns continued the next day. "
                                    else:
                                        narrative += f"Concerns continued over the following {days_span} days, with {last_type} on {last_date_str}. "
                            else:
                                # Non-consecutive - mention most significant other incident
                                most_severe = None
                                for n in notable[1:]:
                                    if n['type'] == 'seclusion':
                                        most_severe = n
                                        break
                                if most_severe is None and len(notable) > 1:
                                    most_severe = notable[-1]

                                if most_severe:
                                    sev_date = most_severe['date'].strftime('%d %B %Y')
                                    sev_reason = most_severe.get('reason', '')
                                    sev_type = linked_incident_type(most_severe['type'], most_severe['date'], most_severe.get('matched', ''), most_severe.get('content_snippet', ''))
                                    if sev_reason and sev_reason in reason_text_map:
                                        narrative += f"A further concern on {sev_date} involving {reason_text_map[sev_reason]} required {sev_type}. "
                                    else:
                                        narrative += f"There was also {sev_type} required on {sev_date}. "

                # Describe medication changes
                if adm_details['medication_changes']:
                    for change_type, meds in adm_details['medication_changes']:
                        if change_type == 'started' and meds:
                            # Apply aggressive deduplication to remove prefix duplicates
                            deduped = dedupe_medication_list(meds)
                            meds_text = format_meds_as_links(meds, limit=3)
                            narrative += f"During the admission, {meds_text} {'was' if len(deduped) == 1 else 'were'} commenced. "
                        # Note: We don't mention discontinued medications - focus on what was given

                # Show discharge medications if available and different from admission
                if adm_details['medications_after']:
                    deduped_discharge = dedupe_medication_list(adm_details['medications_after'])
                    if deduped_discharge:
                        # Check if different from admission medications
                        deduped_admission = dedupe_medication_list(adm_details['medications_before']) if adm_details['medications_before'] else []
                        # Extract base names from dict format
                        def get_base(m):
                            name = m['name'] if isinstance(m, dict) else m
                            match = re.match(r'^([A-Za-z]+)', name)
                            return match.group(1).lower() if match else ''
                        admission_bases = set(get_base(m) for m in deduped_admission)
                        discharge_bases = set(get_base(m) for m in deduped_discharge)

                        # Only mention if there's a difference
                        if discharge_bases != admission_bases or len(deduped_discharge) != len(deduped_admission):
                            meds_text = format_meds_as_links(adm_details['medications_after'], limit=4)
                            narrative += f"On discharge, {pronoun} was on {meds_text}. "

                # Describe improvement and discharge
                improvement_phrase = get_phrase("improvement", pronoun=pronoun, pronoun_poss=pronoun_poss, pronoun_cap=pronoun_cap, pronoun_poss_cap=pronoun_poss_cap)
                discharge_phrase = get_phrase("was_discharged")

                if adm_details['improvement_factors']:
                    # Create links for improvement factors
                    factor_links = []
                    for f in adm_details['improvement_factors'][:2]:
                        if isinstance(f, dict):
                            link = make_link(f['text'], f.get('date'), f['text'], f.get('snippet', ''))
                            factor_links.append(link)
                        else:
                            factor_links.append(f)
                    factors = " and ".join(factor_links)
                    # Add successful leave as outcome if present
                    if adm_details.get('successful_leave'):
                        leave_data = adm_details['successful_leave']
                        if isinstance(leave_data, dict):
                            leave_link = make_link("successful leave", leave_data.get('date'), leave_data.get('matched', 'leave'), leave_data.get('snippet', ''))
                        else:
                            leave_link = "successful leave"
                        narrative += f"{improvement_phrase} with {factors} and took {leave_link}. "
                    else:
                        narrative += f"{improvement_phrase} with {factors}. "
                elif adm_details.get('successful_leave'):
                    leave_data = adm_details['successful_leave']
                    if isinstance(leave_data, dict):
                        leave_link = make_link("successful leave", leave_data.get('date'), leave_data.get('matched', 'leave'), leave_data.get('snippet', ''))
                    else:
                        leave_link = "successful leave"
                    narrative += f"{improvement_phrase} and took {leave_link}. "

                narrative += f"{pronoun_cap} {discharge_phrase} on {fmt_date(adm_end)}."

                # Store discharge date for gap calculation with next admission
                last_discharge_date = adm_end

                # Start new paragraph for post-discharge community period
                narrative += "\n\n"

                # Find the next admission after this discharge (for community details extraction only)
                next_admission = None
                next_admission_date = None
                for next_adm in all_admissions_sorted:
                    next_adm_date = next_adm["date"]
                    if next_adm_date > adm_end:
                        next_admission = next_adm
                        next_admission_date = next_adm_date
                        break

                if next_admission_date:
                    # Calculate time until next admission (for community details extraction)
                    gap_to_next = (next_admission_date - adm_end).days

                    # Gap text is now added at the start of the next admission's description
                    # Here we just extract community details for longer gaps
                    if gap_to_next > 60:
                        # Extract community details for longer community periods
                        # Pass episodes to exclude any inpatient periods within the date range
                        community_details = extract_community_details(all_notes, adm_end, next_admission_date, episodes=episodes)

                        # Describe medications
                        if community_details['medications']:
                            meds_text = format_meds_as_links(community_details['medications'], limit=4)
                            narrative += f"{pronoun_cap} was maintained on {meds_text}. "

                        # Describe contact (who and how)
                        contact_people = community_details.get('contact_people', {})
                        contact_modes = community_details.get('contact_modes', {})

                        if contact_people:
                            # Get main person (highest count)
                            main_person_name = max(contact_people.items(), key=lambda x: x[1]['count'] if isinstance(x[1], dict) else x[1])[0]
                            main_person_data = contact_people[main_person_name]
                            person_link = make_link(main_person_name, main_person_data.get('date') if isinstance(main_person_data, dict) else None, main_person_name, main_person_data.get('snippet', '') if isinstance(main_person_data, dict) else '')
                            if contact_modes:
                                main_mode_name = max(contact_modes.items(), key=lambda x: x[1]['count'] if isinstance(x[1], dict) else x[1])[0]
                                main_mode_data = contact_modes[main_mode_name]
                                mode_link = make_link(main_mode_name, main_mode_data.get('date') if isinstance(main_mode_data, dict) else None, main_mode_name, main_mode_data.get('snippet', '') if isinstance(main_mode_data, dict) else '')
                                narrative += f"During this period {pronoun} had regular {mode_link} contact with {pronoun_poss} {person_link}. "
                            else:
                                narrative += f"During this period {pronoun} had regular contact with {pronoun_poss} {person_link}. "

                        # Describe psychology engagement
                        if community_details['psychology']:
                            psych_links = []
                            seen_types = set()
                            for p in community_details['psychology']:
                                if p['type'] not in seen_types:
                                    seen_types.add(p['type'])
                                    link = make_link(p['type'], p.get('date'), p.get('matched', p['type']), p.get('snippet', ''))
                                    psych_links.append(link)
                            narrative += f"{pronoun_cap} engaged with {', '.join(psych_links)}. "

                        # Describe clinic attendance
                        if community_details['clinics']:
                            clinic_links = []
                            seen_types = set()
                            for c in community_details['clinics'][:2]:
                                if c['type'] not in seen_types:
                                    seen_types.add(c['type'])
                                    link = make_link(c['type'], c.get('date'), c.get('matched', c['type']), c.get('snippet', ''))
                                    clinic_links.append(link)
                            narrative += f"{pronoun_cap} attended {', '.join(clinic_links)}. "

                        # Describe activities
                        if community_details['activities']:
                            activity_links = []
                            for a in community_details['activities'][:3]:
                                if isinstance(a, dict):
                                    link = make_link(a['type'], a.get('date'), a.get('matched', a['type']), a.get('snippet', ''))
                                    activity_links.append(link)
                                else:
                                    activity_links.append(a)
                            narrative += f"{pronoun_cap} engaged in {', '.join(activity_links)}. "

                        # Process crisis events - deduplicate and group by type
                        if community_details['crisis_events']:
                            # Deduplicate by date+type
                            seen = set()
                            unique_events = []
                            for e in community_details['crisis_events']:
                                key = (e['date'], e['type'])
                                if key not in seen:
                                    seen.add(key)
                                    unique_events.append(e)

                            # Group by type and collect events (with full info for linking)
                            events_by_type = {}
                            for e in unique_events:
                                etype = e['type']
                                if etype not in events_by_type:
                                    events_by_type[etype] = []
                                events_by_type[etype].append(e)

                            # Sort by date within each type and format output with links
                            crisis_parts = []
                            for etype, events in events_by_type.items():
                                events = sorted(events, key=lambda x: x['date'])
                                first_event = events[0]
                                etype_link = make_link(etype, first_event['date'], etype, first_event.get('summary', '')[:100])
                                dates = [e['date'] for e in events]
                                if len(dates) == 1:
                                    date_str = dates[0].strftime('%d %B %Y')
                                    crisis_parts.append(f"on {date_str} {etype_link} were involved")
                                elif len(dates) == 2:
                                    # Check if consecutive days
                                    if (dates[1] - dates[0]).days == 1:
                                        date_str = f"{dates[0].strftime('%d')} and {dates[1].strftime('%d %B %Y')}"
                                    else:
                                        date_str = f"{dates[0].strftime('%d %B')} and {dates[1].strftime('%d %B %Y')}"
                                    crisis_parts.append(f"on {date_str} {etype_link} were involved")
                                else:
                                    # Multiple dates - just mention first and last
                                    date_str = f"between {dates[0].strftime('%d %B')} and {dates[-1].strftime('%d %B %Y')}"
                                    crisis_parts.append(f"{date_str} {etype_link} were involved")

                            if crisis_parts:
                                narrative += f"{crisis_parts[0].capitalize()}"
                                if len(crisis_parts) > 1:
                                    narrative += f" and {crisis_parts[1]}"
                                narrative += ". "

                        # Add serious concerns with dates
                        dated_events = []
                        if community_details['concerns']:
                            serious_concerns = ['safeguarding', 'substance use']
                            for c in community_details['concerns']:
                                if c['type'] in serious_concerns:
                                    dated_events.append({
                                        'date': c['date'],
                                        'text': f"{c['type']} concerns were raised"
                                    })

                        # Sort by date and output concerns
                        if dated_events:
                            dated_events.sort(key=lambda x: x['date'])
                            for event in dated_events[:2]:  # Limit to 2 concern events
                                date_str = event['date'].strftime('%d %B %Y')
                                narrative += f"On {date_str}, {event['text']}. "

                        # Other concerns summarised (without dates)
                        if community_details['concerns']:
                            # Filter out safeguarding, substance use, and any type containing 'concern' or 'noted'
                            excluded = ['safeguarding', 'substance use']
                            other_found = [c for c in community_details['concerns']
                                          if c['type'] not in excluded
                                          and 'concern' not in c['type'].lower()
                                          and 'noted' not in c['type'].lower()]
                            other_types = list(set(c['type'] for c in other_found))
                            if other_types:
                                narrative += f"Concerns were also noted around {', '.join(other_types[:2])}. "

                        # Describe incidents during community period - consolidate by type
                        if community_details.get('incidents'):
                            incidents_by_type = defaultdict(list)
                            for inc in community_details['incidents']:
                                incidents_by_type[inc['type']].append(inc)

                            for inc_type, inc_list in incidents_by_type.items():
                                if len(inc_list) == 1:
                                    inc = inc_list[0]
                                    date_str = inc['date'].strftime('%d %B %Y')
                                    inc_link = make_link(inc_type, inc['date'], inc_type, inc.get('content_snippet', ''))
                                    narrative += f"On {date_str}, there was {inc_link}. "
                                else:
                                    # Multiple incidents of same type - consolidate
                                    sorted_incs = sorted(inc_list, key=lambda x: x['date'])
                                    first_inc = sorted_incs[0]
                                    last_inc = sorted_incs[-1]
                                    first_date = first_inc['date']
                                    last_date = last_inc['date']
                                    inc_link = make_link(inc_type, first_inc['date'], inc_type, first_inc.get('content_snippet', ''))

                                    if first_date.month == last_date.month and first_date.year == last_date.year:
                                        date_range = f"between {first_date.day}-{last_date.day} {first_date.strftime('%B %Y')}"
                                    elif first_date.year == last_date.year:
                                        date_range = f"between {first_date.strftime('%d %B')} and {last_date.strftime('%d %B %Y')}"
                                    else:
                                        date_range = f"between {first_date.strftime('%d %B %Y')} and {last_date.strftime('%d %B %Y')}"

                                    narrative += f"There were {len(inc_list)} instances of {inc_link} {date_range}. "

                        # Describe substance misuse
                        if community_details.get('substance_misuse'):
                            substances = list(set(s['type'] for s in community_details['substance_misuse']))
                            if substances:
                                first_substance = community_details['substance_misuse'][0]
                                substance_link = make_link(', '.join(substances[:2]), first_substance['date'], first_substance['type'], first_substance.get('content_snippet', ''))
                                narrative += f"There were concerns about {substance_link} use during this period. "

                        # Note: AWOL/absconding removed - not applicable to community periods

                else:
                    # No further admissions - describe ongoing community care in detail
                    following_phrase = get_phrase("following_discharge")

                    # Calculate time in community from discharge to now
                    from_date = adm_end
                    to_date = datetime.now().date() if hasattr(datetime.now(), 'date') else datetime.now()

                    # Pass episodes to exclude any inpatient periods within the date range
                    community_details = extract_community_details(all_notes, from_date, to_date, episodes=episodes)

                    narrative += f"{following_phrase}, {pronoun} has remained in the community. "

                    # Describe medications
                    if community_details['medications']:
                        meds_text = format_meds_as_links(community_details['medications'], limit=4)
                        narrative += f"{pronoun_cap} is currently maintained on {meds_text}. "

                    # Describe contact (who and how)
                    contact_people = community_details.get('contact_people', {})
                    contact_modes = community_details.get('contact_modes', {})

                    if contact_people:
                        # Get main person (highest count)
                        main_person_name = max(contact_people.items(), key=lambda x: x[1]['count'] if isinstance(x[1], dict) else x[1])[0]
                        main_person_data = contact_people[main_person_name]
                        person_link = make_link(main_person_name, main_person_data.get('date') if isinstance(main_person_data, dict) else None, main_person_name, main_person_data.get('snippet', '') if isinstance(main_person_data, dict) else '')
                        if contact_modes:
                            main_mode_name = max(contact_modes.items(), key=lambda x: x[1]['count'] if isinstance(x[1], dict) else x[1])[0]
                            main_mode_data = contact_modes[main_mode_name]
                            mode_link = make_link(main_mode_name, main_mode_data.get('date') if isinstance(main_mode_data, dict) else None, main_mode_name, main_mode_data.get('snippet', '') if isinstance(main_mode_data, dict) else '')
                            narrative += f"{pronoun_cap} has ongoing {mode_link} contact with {pronoun_poss} {person_link}. "
                        else:
                            narrative += f"{pronoun_cap} has ongoing contact with {pronoun_poss} {person_link}. "

                    # Describe psychology engagement
                    if community_details['psychology']:
                        psych_links = []
                        seen_types = set()
                        for p in community_details['psychology']:
                            if p['type'] not in seen_types:
                                seen_types.add(p['type'])
                                link = make_link(p['type'], p.get('date'), p.get('matched', p['type']), p.get('snippet', ''))
                                psych_links.append(link)
                        narrative += f"{pronoun_cap} has engaged with {', '.join(psych_links)}. "

                    # Describe clinic attendance
                    if community_details['clinics']:
                        clinic_links = []
                        seen_types = set()
                        for c in community_details['clinics'][:2]:
                            if c['type'] not in seen_types:
                                seen_types.add(c['type'])
                                link = make_link(c['type'], c.get('date'), c.get('matched', c['type']), c.get('snippet', ''))
                                clinic_links.append(link)
                        narrative += f"{pronoun_cap} attends {', '.join(clinic_links)}. "

                    # Describe activities
                    if community_details['activities']:
                        activity_links = []
                        for a in community_details['activities'][:3]:
                            if isinstance(a, dict):
                                link = make_link(a['type'], a.get('date'), a.get('matched', a['type']), a.get('snippet', ''))
                                activity_links.append(link)
                            else:
                                activity_links.append(a)
                        narrative += f"{pronoun_cap} has engaged in {', '.join(activity_links)}. "

                    # Process crisis events - deduplicate and group by type
                    if community_details['crisis_events']:
                        # Deduplicate by date+type
                        seen = set()
                        unique_events = []
                        for e in community_details['crisis_events']:
                            key = (e['date'], e['type'])
                            if key not in seen:
                                seen.add(key)
                                unique_events.append(e)

                        # Group by type and collect events (with full info for linking)
                        events_by_type = {}
                        for e in unique_events:
                            etype = e['type']
                            if etype not in events_by_type:
                                events_by_type[etype] = []
                            events_by_type[etype].append(e)

                        # Sort by date within each type and format output with links
                        crisis_parts = []
                        for etype, events in events_by_type.items():
                            events = sorted(events, key=lambda x: x['date'])
                            first_event = events[0]
                            etype_link = make_link(etype, first_event['date'], etype, first_event.get('summary', '')[:100])
                            dates = [e['date'] for e in events]
                            if len(dates) == 1:
                                date_str = dates[0].strftime('%d %B %Y')
                                crisis_parts.append(f"on {date_str} {etype_link} were involved")
                            elif len(dates) == 2:
                                # Check if consecutive days
                                if (dates[1] - dates[0]).days == 1:
                                    date_str = f"{dates[0].strftime('%d')} and {dates[1].strftime('%d %B %Y')}"
                                else:
                                    date_str = f"{dates[0].strftime('%d %B')} and {dates[1].strftime('%d %B %Y')}"
                                crisis_parts.append(f"on {date_str} {etype_link} were involved")
                            else:
                                # Multiple dates - just mention first and last
                                date_str = f"between {dates[0].strftime('%d %B')} and {dates[-1].strftime('%d %B %Y')}"
                                crisis_parts.append(f"{date_str} {etype_link} were involved")

                        if crisis_parts:
                            narrative += f"{crisis_parts[0].capitalize()}"
                            if len(crisis_parts) > 1:
                                narrative += f" and {crisis_parts[1]}"
                            narrative += ". "

                    # Add serious concerns with dates
                    dated_events = []
                    if community_details['concerns']:
                        serious_concerns = ['safeguarding', 'substance use']
                        for c in community_details['concerns']:
                            if c['type'] in serious_concerns:
                                dated_events.append({
                                    'date': c['date'],
                                    'text': f"{c['type']} concerns were raised"
                                })

                    # Sort by date and output concerns
                    if dated_events:
                        dated_events.sort(key=lambda x: x['date'])
                        for event in dated_events[:2]:  # Limit to 2 concern events
                            date_str = event['date'].strftime('%d %B %Y')
                            narrative += f"On {date_str}, {event['text']}. "

                    # Other concerns summarised (without dates)
                    if community_details['concerns']:
                        # Filter out safeguarding, substance use, and any type containing 'concern' or 'noted'
                        excluded = ['safeguarding', 'substance use']
                        other_found = [c for c in community_details['concerns']
                                      if c['type'] not in excluded
                                      and 'concern' not in c['type'].lower()
                                      and 'noted' not in c['type'].lower()]
                        other_types = list(set(c['type'] for c in other_found))
                        if other_types:
                            narrative += f"Concerns have also been noted around {', '.join(other_types[:2])}. "

                    # Describe incidents during current community period - consolidate by type
                    if community_details.get('incidents'):
                        incidents_by_type = defaultdict(list)
                        for inc in community_details['incidents']:
                            incidents_by_type[inc['type']].append(inc)

                        for inc_type, inc_list in incidents_by_type.items():
                            if len(inc_list) == 1:
                                inc = inc_list[0]
                                date_str = inc['date'].strftime('%d %B %Y')
                                inc_link = make_link(inc_type, inc['date'], inc_type, inc.get('content_snippet', ''))
                                narrative += f"On {date_str}, there was {inc_link}. "
                            else:
                                # Multiple incidents of same type - consolidate
                                sorted_incs = sorted(inc_list, key=lambda x: x['date'])
                                first_inc = sorted_incs[0]
                                last_inc = sorted_incs[-1]
                                first_date = first_inc['date']
                                last_date = last_inc['date']
                                inc_link = make_link(inc_type, first_inc['date'], inc_type, first_inc.get('content_snippet', ''))

                                if first_date.month == last_date.month and first_date.year == last_date.year:
                                    date_range = f"between {first_date.day}-{last_date.day} {first_date.strftime('%B %Y')}"
                                elif first_date.year == last_date.year:
                                    date_range = f"between {first_date.strftime('%d %B')} and {last_date.strftime('%d %B %Y')}"
                                else:
                                    date_range = f"between {first_date.strftime('%d %B %Y')} and {last_date.strftime('%d %B %Y')}"

                                narrative += f"There have been {len(inc_list)} instances of {inc_link} {date_range}. "

                    # Describe substance misuse
                    if community_details.get('substance_misuse'):
                        substances = list(set(s['type'] for s in community_details['substance_misuse']))
                        if substances:
                            first_substance = community_details['substance_misuse'][0]
                            substance_link = make_link(', '.join(substances[:2]), first_substance['date'], first_substance['type'], first_substance.get('content_snippet', ''))
                            narrative += f"There have been concerns about {substance_link} use. "

                    # Note: AWOL/absconding removed - not applicable to community periods

            else:
                # No end date - admission ongoing or unknown
                admission_phrase = get_phrase("required_admission")
                narrative += f"{pronoun_cap} {admission_phrase} on {fmt_date(adm_date)}. "

            last_mentioned_state = "admission"

        # Note: Discharges are now handled within the admission loop above
        # No need for standalone discharge handling as it causes out-of-order statements

        # If first group with incidents, describe the community management
        if group_idx == 0 and year_incidents and not year_admissions:
            # Check if managed in community (no admission that year)
            narrative += f"These were managed by the community team without requiring admission. "
            last_mentioned_state = "stable"

    # Add summary conclusion
    narrative += f"""

SUMMARY:
"""
    if total_incidents == 0:
        narrative += f"{name} demonstrated a stable presentation throughout the review period with no significant risk concerns documented.\n"
    elif total_violence == 0 and total_verbal > 0:
        narrative += f"While {total_verbal} verbal aggression concern{'s were' if total_verbal > 1 else ' was'} recorded, there were no episodes of physical violence during the review period.\n"
    elif total_violence > 0:
        narrative += f"The review period included {total_violence} physical aggression concern{'s' if total_violence > 1 else ''} and {total_verbal} verbal aggression concern{'s' if total_verbal > 1 else ''}.\n"

    if len(inpatient_episodes) > 0:
        narrative += f"{pronoun_cap} had {len(inpatient_episodes)} admission{'s' if len(inpatient_episodes) > 1 else ''} totalling {total_inpatient_days} inpatient days.\n"

    narrative += f"""
---
Report generated: {datetime.now().strftime('%d %B %Y')}
"""
    # Post-process to consolidate repetitive statements
    narrative = consolidate_narrative(narrative)
    return narrative


def consolidate_narrative(narrative: str) -> str:
    """
    Second-pass consolidation of the narrative to remove repetitive patterns.

    Identifies patterns like:
    - "On 17 July 2012, there was police involvement. On 16 July 2012, there was police involvement."

    And consolidates them into:
    - "Between 13-17 July 2012, there were multiple instances of police involvement."
    """
    from collections import defaultdict

    # Pattern to match "On [date], there was [event]." or "On [date], [pronoun] [verb] [event]."
    # Captures: full match, day, month, year, event description
    # Updated to properly handle linked text like <a href="...">police involvement</a>
    date_event_pattern = re.compile(
        r'On (\d{1,2}(?:st|nd|rd|th)?) ([A-Z][a-z]+) (\d{4}), (there was |there were |[Hh]e |[Ss]he |[Tt]hey )?(an? )?(<a[^>]*>[^<]+</a>|[^.<]+)\. ?',
        re.IGNORECASE
    )

    # Find all date-event statements
    matches = list(date_event_pattern.finditer(narrative))

    if len(matches) < 2:
        return narrative

    # Group by event type (normalized)
    event_groups = defaultdict(list)

    for match in matches:
        day = match.group(1).rstrip('stndrdth')
        month = match.group(2)
        year = match.group(3)
        prefix = match.group(4) or ''
        article = match.group(5) or ''
        event = match.group(6).strip()

        # Normalize event text for grouping (strip HTML tags for comparison)
        event_normalized = re.sub(r'<[^>]+>', '', event).lower().strip()

        # Create date object for sorting
        try:
            month_num = datetime.strptime(month, '%B').month
            date_obj = datetime(int(year), month_num, int(day))
        except:
            continue

        event_groups[event_normalized].append({
            'date': date_obj,
            'day': day,
            'month': month,
            'year': year,
            'event': event,  # Keep original with links
            'full_match': match.group(0),
            'prefix': prefix,
            'article': article,
            'start': match.start(),
            'end': match.end()
        })

    # Find groups with 2+ repetitions that are close together
    consolidations = []

    for event_type, occurrences in event_groups.items():
        if len(occurrences) < 2:
            continue

        # Sort by date
        occurrences.sort(key=lambda x: x['date'])

        # Check if dates are within a reasonable range
        first_date = occurrences[0]['date']
        last_date = occurrences[-1]['date']
        date_span = (last_date - first_date).days

        # Same-date duplicates: just keep one (remove the rest)
        if date_span == 0 and len(occurrences) >= 2:
            consolidations.append({
                'event_type': event_type,
                'occurrences': occurrences,
                'first_date': first_date,
                'last_date': last_date,
                'count': len(occurrences),
                'same_date': True  # Flag for same-date handling
            })
            continue

        # Only consolidate if within 60 days
        # For 2 occurrences, be stricter (within 14 days)
        # For 3+ occurrences, allow up to 60 days
        if len(occurrences) == 2 and date_span <= 14:
            consolidations.append({
                'event_type': event_type,
                'occurrences': occurrences,
                'first_date': first_date,
                'last_date': last_date,
                'count': len(occurrences)
            })
        elif len(occurrences) >= 3 and date_span <= 60:
            consolidations.append({
                'event_type': event_type,
                'occurrences': occurrences,
                'first_date': first_date,
                'last_date': last_date,
                'count': len(occurrences)
            })

    if not consolidations:
        return narrative

    # Sort consolidations by position (to process from end to start to preserve positions)
    for cons in consolidations:
        cons['first_pos'] = min(o['start'] for o in cons['occurrences'])
    consolidations.sort(key=lambda x: x['first_pos'], reverse=True)

    # Apply consolidations
    for cons in consolidations:
        occurrences = cons['occurrences']
        first = occurrences[0]
        last = occurrences[-1]
        count = cons['count']

        # Handle same-date duplicates: just keep the first one
        if cons.get('same_date'):
            # Remove all but the first occurrence
            occurrences_sorted = sorted(occurrences, key=lambda x: x['start'], reverse=True)
            for i, occ in enumerate(occurrences_sorted):
                if i < len(occurrences_sorted) - 1:
                    # Remove this duplicate (keep the first one which is last in reversed list)
                    narrative = narrative[:occ['start']] + narrative[occ['end']:]
            continue

        # Build consolidated statement for different dates
        if first['month'] == last['month'] and first['year'] == last['year']:
            # Same month: "Between 13-17 July 2012"
            date_range = f"between {first['day']}-{last['day']} {first['month']} {first['year']}"
        elif first['year'] == last['year']:
            # Same year, different months
            date_range = f"between {first['day']} {first['month']} and {last['day']} {last['month']} {first['year']}"
        else:
            # Different years
            date_range = f"between {first['day']} {first['month']} {first['year']} and {last['day']} {last['month']} {last['year']}"

        # Use the first occurrence's event text (preserves link)
        event_text = first['event']

        # Create consolidated statement
        consolidated = f"There were {count} instances of {event_text} {date_range}. "

        # Remove all individual occurrences and insert consolidated at first position
        # Process from end to start
        occurrences_sorted = sorted(occurrences, key=lambda x: x['start'], reverse=True)

        for i, occ in enumerate(occurrences_sorted):
            if i == len(occurrences_sorted) - 1:
                # Last one (first in original order) - replace with consolidated
                narrative = narrative[:occ['start']] + consolidated + narrative[occ['end']:]
            else:
                # Remove this occurrence
                narrative = narrative[:occ['start']] + narrative[occ['end']:]

    # Clean up any double spaces or orphaned punctuation
    narrative = re.sub(r'  +', ' ', narrative)
    narrative = re.sub(r'\. +\.', '.', narrative)
    narrative = re.sub(r' +\.', '.', narrative)

    return narrative


def create_progress_timeline_chart(results: Dict) -> tuple:
    """Create compact risk level timeline with markers on a single row below.

    Modeled after create_risk_timeline_visual from risk_overview_panel.
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
        import matplotlib.lines as mlines
        from io import BytesIO

        all_months = results.get("all_months", [])
        monthly_data = results.get("monthly_data", {})
        tentpole_events = results.get("tentpole_events", [])
        monthly_violence = results.get("monthly_violence", {})
        monthly_verbal = results.get("monthly_verbal", {})
        monthly_incidents = results.get("monthly_incidents", {})

        if not all_months or len(all_months) < 2:
            return None, {}

        # Calculate risk level for each month
        month_data = []
        for month in all_months:
            data = monthly_data.get(month, {})
            total = data.get("total_incidents", 0)

            # Also count from monthly_incidents if not in monthly_data
            if total == 0:
                total = monthly_incidents.get(month, 0)

            # Categorize activity level (same as risk_overview_panel)
            if total == 0:
                level = 0
                color = '#2d5a3d'  # dark green
                level_name = "Quiet"
            elif total <= 3:
                level = 1
                color = '#22c55e'  # green
                level_name = "Low"
            elif total <= 8:
                level = 2
                color = '#f59e0b'  # amber
                level_name = "Moderate"
            elif total <= 15:
                level = 3
                color = '#f97316'  # orange
                level_name = "Elevated"
            else:
                level = 4
                color = '#ef4444'  # red
                level_name = "High"

            month_data.append({
                "month": month,
                "level": level,
                "level_name": level_name,
                "color": color,
                "total": total,
            })

        num_months = len(all_months)
        has_markers = bool(monthly_violence) or bool(monthly_verbal) or bool(tentpole_events)

        # Compact figure dimensions
        fig_width = max(12, num_months * 0.4)
        fig_height = 2.8 if has_markers else 2.2

        dpi = 100
        fig, ax = plt.subplots(figsize=(fig_width, fig_height))
        fig.patch.set_facecolor('#1a1a2e')
        ax.set_facecolor('#1a1a2e')

        # Y positions - compact layout
        bar_height = 0.6
        risk_y = 0.5
        markers_y = -0.15  # Single row for all markers, just below risk bar

        # Draw risk level bars
        for i, data in enumerate(month_data):
            ax.barh(risk_y, 1, left=i, height=bar_height, color=data["color"],
                   edgecolor='#333', linewidth=0.5)

            # Add incident count on bar if > 0
            if data["total"] > 0:
                ax.text(i + 0.5, risk_y, str(data["total"]), ha='center', va='center',
                       fontsize=8, color='white', fontweight='bold')

        # Prepare marker positions for interactivity
        month_to_pos = {m: i for i, m in enumerate(all_months)}
        event_markers = []
        violence_markers = []
        verbal_markers = []

        # Key Events (diamonds) - on shared marker row
        if tentpole_events:
            type_priority = {'Tribunal': 1, 'Managers Hearing': 2, 'Unescorted Leave': 3,
                           'Overnight Leave': 4, 'Community Leave': 5, 'Ground Leave': 6}
            type_labels = {
                'Tribunal': 'T', 'Managers Hearing': 'M', 'Unescorted Leave': 'U',
                'Overnight Leave': 'O', 'Community Leave': 'C', 'Ground Leave': 'G',
                'Escorted Leave': 'E', 'CPA Review': 'P', 'Ward Round': 'W',
                'Section Change': 'S', 'Medication Change': 'X',
            }

            # Group events by month, keep highest priority
            monthly_events = {}
            for event in tentpole_events:
                event_month = event["month"]
                event_type = event["type"]
                priority = type_priority.get(event_type, 10)
                if event_month not in monthly_events or priority < monthly_events[event_month][1]:
                    monthly_events[event_month] = (event, priority, event_type)

            for event_month, (event, _, event_type) in monthly_events.items():
                if event_month in month_to_pos:
                    pos = month_to_pos[event_month]
                    # Offset slightly up within marker row
                    ax.plot(pos + 0.5, markers_y + 0.08, 'D', markersize=7, color='#64B5F6', zorder=5)
                    event_markers.append({
                        'month': event_month,
                        'pos': pos,
                        'event': event,
                    })

        # Violence (triangles) - on shared marker row
        for month, count in monthly_violence.items():
            if month in month_to_pos and count > 0:
                pos = month_to_pos[month]
                marker_size = min(6 + count * 2, 12)
                # Offset slightly down within marker row
                ax.plot(pos + 0.5, markers_y - 0.08, '^', markersize=marker_size,
                       color='#EF5350', zorder=5)
                violence_markers.append({
                    'month': month,
                    'pos': pos,
                    'count': count,
                })

        # Verbal (circles) - on shared marker row
        for month, count in monthly_verbal.items():
            if month in month_to_pos and count > 0:
                pos = month_to_pos[month]
                size = min(count * 25, 80)
                # Center in marker row
                ax.scatter(pos + 0.5, markers_y, s=size, alpha=0.7,
                          color='#AB47BC', zorder=4)
                verbal_markers.append({
                    'month': month,
                    'pos': pos,
                    'count': count,
                })

        # X-axis labels - directly below markers, show more months
        ax.set_xlim(0, num_months)
        ax.set_ylim(-0.5 if has_markers else 0, 1.1)

        # Show more months on x-axis (every 2-3 months instead of sparse)
        if num_months <= 6:
            step = 1
        elif num_months <= 12:
            step = 1
        elif num_months <= 24:
            step = 2
        elif num_months <= 48:
            step = 3
        else:
            step = max(1, num_months // 20)  # Show ~20 labels max

        tick_positions = list(range(0, num_months, step))
        tick_labels = []
        for i in tick_positions:
            try:
                dt = datetime.strptime(all_months[i], "%Y-%m")
                tick_labels.append(dt.strftime("%b '%y"))
            except:
                tick_labels.append(all_months[i])

        ax.set_xticks([p + 0.5 for p in tick_positions])
        ax.set_xticklabels(tick_labels, fontsize=8, color='#AAA', rotation=45, ha='right')

        # Remove y-axis ticks
        ax.set_yticks([])

        # Title
        ax.set_title('Progress Timeline', fontsize=11, color='white', pad=8, loc='left')

        # Combined legend - risk levels + marker types
        legend_elements = [
            mpatches.Patch(facecolor='#2d5a3d', edgecolor='#555', label='Quiet'),
            mpatches.Patch(facecolor='#22c55e', edgecolor='#555', label='Low'),
            mpatches.Patch(facecolor='#f59e0b', edgecolor='#555', label='Moderate'),
            mpatches.Patch(facecolor='#f97316', edgecolor='#555', label='Elevated'),
            mpatches.Patch(facecolor='#ef4444', edgecolor='#555', label='High'),
        ]
        # Add marker legend items
        if tentpole_events:
            legend_elements.append(mlines.Line2D([], [], marker='D', color='#64B5F6',
                                   linestyle='None', markersize=6, label='Event'))
        if monthly_violence:
            legend_elements.append(mlines.Line2D([], [], marker='^', color='#EF5350',
                                   linestyle='None', markersize=6, label='Violence'))
        if monthly_verbal:
            legend_elements.append(mlines.Line2D([], [], marker='o', color='#AB47BC',
                                   linestyle='None', markersize=6, label='Verbal'))

        ax.legend(handles=legend_elements, loc='upper right', fontsize=7,
                 facecolor='#2a2a3e', edgecolor='#555', labelcolor='white',
                 ncol=min(len(legend_elements), 8), framealpha=0.9)

        # Remove spines
        for spine in ax.spines.values():
            spine.set_visible(False)

        plt.tight_layout()

        # Save to bytes
        buf = BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, facecolor='#1a1a2e',
                   edgecolor='none', bbox_inches='tight', pad_inches=0.1)

        # Get actual image dimensions for coordinate mapping
        buf.seek(0)
        from PIL import Image
        img = Image.open(buf)
        img_width, img_height = img.size
        buf.seek(0)

        # Calculate bar positions for interactivity
        left_margin = 50
        right_margin = 30
        top_margin = 40
        bottom_margin = 45

        chart_width = img_width - left_margin - right_margin
        bar_width = chart_width / num_months

        # Map y positions to pixel coordinates
        y_max = 1.1
        y_min = -0.5 if has_markers else 0
        y_range = y_max - y_min
        chart_height = img_height - top_margin - bottom_margin

        def y_to_pixel(y_data):
            normalized = (y_max - y_data) / y_range
            return top_margin + normalized * chart_height

        # Risk bar positions
        risk_y_top = y_to_pixel(risk_y + bar_height/2)
        risk_y_bottom = y_to_pixel(risk_y - bar_height/2)

        all_bar_positions = []

        for i, data in enumerate(month_data):
            x_left = left_margin + i * bar_width
            x_right = left_margin + (i + 1) * bar_width

            all_bar_positions.append({
                'type': 'risk',
                'index': i,
                'month': data['month'],
                'month_label': _get_month_name(data['month']),
                'x_left': x_left,
                'x_right': x_right,
                'y_top': risk_y_top,
                'y_bottom': risk_y_bottom,
                'total': data['total'],
                'level_name': data['level_name'],
            })

        # Marker positions (all on same row)
        markers_y_pixel = y_to_pixel(markers_y)

        for marker in event_markers:
            x_center = left_margin + (marker['pos'] + 0.5) * bar_width
            all_bar_positions.append({
                'type': 'event',
                'month': marker['month'],
                'month_label': _get_month_name(marker['month']),
                'x_left': x_center - 15,
                'x_right': x_center + 15,
                'y_top': markers_y_pixel - 15,
                'y_bottom': markers_y_pixel + 15,
                'event': marker['event'],
            })

        for marker in violence_markers:
            x_center = left_margin + (marker['pos'] + 0.5) * bar_width
            all_bar_positions.append({
                'type': 'violence',
                'month': marker['month'],
                'month_label': _get_month_name(marker['month']),
                'x_left': x_center - 15,
                'x_right': x_center + 15,
                'y_top': markers_y_pixel - 15,
                'y_bottom': markers_y_pixel + 15,
                'count': marker['count'],
            })

        for marker in verbal_markers:
            x_center = left_margin + (marker['pos'] + 0.5) * bar_width
            all_bar_positions.append({
                'type': 'verbal',
                'month': marker['month'],
                'month_label': _get_month_name(marker['month']),
                'x_left': x_center - 15,
                'x_right': x_center + 15,
                'y_top': markers_y_pixel - 15,
                'y_bottom': markers_y_pixel + 15,
                'count': marker['count'],
            })

        plt.close(fig)

        timeline_info = {
            'bar_positions': all_bar_positions,
            'months': all_months,
            'month_data': month_data,
        }

        return buf.read(), timeline_info

    except Exception as e:
        print(f"Error creating progress timeline: {e}")
        import traceback
        traceback.print_exc()
        return None, {}


# ============================================================
# INTERACTIVE TIMELINE - With Tooltips (Risk Level Style)
# ============================================================

class InteractiveProgressTimeline(QLabel):
    """Timeline with hover tooltips and click-to-scroll."""

    eventSelected = Signal(object)

    def __init__(self):
        super().__init__()
        self.setMouseTracking(True)
        self.setCursor(Qt.PointingHandCursor)
        self.setAttribute(Qt.WA_Hover, True)
        self.timeline_info = {}
        self.monthly_data = {}  # Store incident data for popups

    def set_chart(self, image_bytes: bytes, timeline_info: dict, monthly_data: dict = None):
        """Set the chart image and metadata."""
        self.timeline_info = timeline_info
        self.monthly_data = monthly_data or {}
        if image_bytes:
            pixmap = QPixmap()
            pixmap.loadFromData(image_bytes)
            self.setPixmap(pixmap)

    def _get_bar_at_pos(self, pos):
        """Return bar info if position is over a bar."""
        if not self.timeline_info.get('bar_positions'):
            return None

        x, y = pos.x(), pos.y()

        for bar in self.timeline_info['bar_positions']:
            if bar['x_left'] <= x <= bar['x_right'] and bar['y_top'] <= y <= bar['y_bottom']:
                return bar

        return None

    def mouseMoveEvent(self, event):
        """Show tooltip on hover."""
        bar = self._get_bar_at_pos(event.pos())

        if bar:
            bar_type = bar.get('type', 'risk')
            month_label = bar.get('month_label', bar.get('month', ''))

            if bar_type == 'risk':
                level_name = bar.get('level_name', '')
                total = bar.get('total', 0)
                tooltip = f"<b>{month_label}</b><br>"
                tooltip += f"Risk Level: <b>{level_name}</b><br>"
                tooltip += f"Total Incidents: <b>{total}</b><br>"
                if total > 0:
                    tooltip += "<br><i>Click to view incidents</i>"

            elif bar_type == 'event':
                event_data = bar.get('event', {})
                date = event_data.get('date')
                date_str = date.strftime("%d %b %Y") if date else month_label
                tooltip = f"<b>{event_data.get('type', 'Event')}</b><br>"
                tooltip += f"<b>Date:</b> {date_str}<br><br>"
                tooltip += "<i>Click to view in notes</i>"

            elif bar_type == 'violence':
                count = bar.get('count', 0)
                tooltip = f"<b>Violence Incidents</b><br>"
                tooltip += f"<b>{month_label}:</b> {count} incident{'s' if count != 1 else ''}<br><br>"
                tooltip += "<i>Click to view incidents</i>"

            elif bar_type == 'verbal':
                count = bar.get('count', 0)
                tooltip = f"<b>Verbal Aggression</b><br>"
                tooltip += f"<b>{month_label}:</b> {count} incident{'s' if count != 1 else ''}<br><br>"
                tooltip += "<i>Click to view incidents</i>"

            else:
                tooltip = f"<b>{month_label}</b>"

            # Wrap in HTML with explicit styling for Windows
            tooltip = f"<div style='background-color: #fffbe6; color: #000000; padding: 4px;'>{tooltip}</div>"
            QToolTip.showText(event.globalPosition().toPoint(), tooltip, self)
        else:
            QToolTip.hideText()

        super().mouseMoveEvent(event)

    def mousePressEvent(self, event):
        """Show popup menu with incidents when clicked."""
        if event.button() == Qt.LeftButton:
            bar = self._get_bar_at_pos(event.pos())
            if bar:
                bar_type = bar.get('type', '')
                month = bar.get('month')
                month_label = bar.get('month_label', month)

                # For events, emit to scroll to note with event type for highlighting
                if bar_type == 'event' and bar.get('event'):
                    evt = bar['event']
                    # Add matched text for highlighting (use event type)
                    evt_with_match = dict(evt)
                    evt_with_match['matched'] = evt.get('type', '')
                    self.eventSelected.emit(evt_with_match)
                    super().mousePressEvent(event)
                    return

                # For risk/violence/verbal, show popup with incidents
                if month and month in self.monthly_data:
                    data = self.monthly_data[month]
                    incidents = data.get('incidents', [])
                    events = data.get('tentpole_events', [])

                    # Filter by type if specific marker clicked
                    if bar_type == 'violence':
                        incidents = [i for i in incidents if i.get('category') == 'Physical Violence']
                    elif bar_type == 'verbal':
                        incidents = [i for i in incidents if i.get('category') == 'Verbal Aggression']

                    if incidents:
                        self._show_incidents_popup(event, month_label, bar, incidents, events)
                    elif bar_type in ('risk', 'violence', 'verbal'):
                        # Still emit for scrolling even if no detailed incidents
                        try:
                            dt = datetime.strptime(month, "%Y-%m")
                            self.eventSelected.emit({'date': dt, 'type': bar_type})
                        except:
                            pass

        super().mousePressEvent(event)

    def _show_incidents_popup(self, event, month_label, bar, incidents, events):
        """Show scrollable popup with incidents for this month."""
        if not incidents:
            return

        # Create popup widget
        popup = QWidget(self, Qt.Popup | Qt.FramelessWindowHint)
        popup.setAttribute(Qt.WA_DeleteOnClose)
        popup.setStyleSheet("""
            QWidget {
                background-color: rgba(35,35,45,0.98);
                border: 1px solid #555;
                border-radius: 8px;
            }
        """)

        layout = QVBoxLayout(popup)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        # Header
        bar_type = bar.get('type', 'risk')
        level_name = bar.get('level_name', '')

        if bar_type == 'risk':
            header_text = f"{month_label} - {level_name} ({len(incidents)} incidents)"
        elif bar_type == 'violence':
            header_text = f"{month_label} - Violence ({len(incidents)} incidents)"
        elif bar_type == 'verbal':
            header_text = f"{month_label} - Verbal ({len(incidents)} incidents)"
        else:
            header_text = f"{month_label} ({len(incidents)} incidents)"

        header = QLabel(header_text)
        header.setStyleSheet("font-weight: bold; font-size: 13px; color: #FFF; padding: 4px;")
        layout.addWidget(header)

        # Separator
        sep = QFrame()
        sep.setFrameShape(QFrame.HLine)
        sep.setStyleSheet("background-color: #555;")
        sep.setFixedHeight(1)
        layout.addWidget(sep)

        # Scrollable list area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setStyleSheet("""
            QScrollArea {
                background: transparent;
                border: none;
            }
            QScrollBar:vertical {
                background: rgba(255,255,255,0.1);
                width: 8px;
                border-radius: 4px;
            }
            QScrollBar::handle:vertical {
                background: rgba(255,255,255,0.3);
                border-radius: 4px;
                min-height: 30px;
            }
            QScrollBar::handle:vertical:hover {
                background: rgba(255,255,255,0.5);
            }
        """)
        scroll.viewport().setStyleSheet("background: transparent;")

        list_widget = QWidget()
        list_widget.setStyleSheet("background: transparent;")
        list_layout = QVBoxLayout(list_widget)
        list_layout.setContentsMargins(0, 0, 0, 0)
        list_layout.setSpacing(2)

        # Store reference for click handling
        popup._selected_data = None

        def make_item_click_handler(data):
            def handler():
                popup._selected_data = data
                popup.close()
            return handler

        # Add incidents only (no events)
        for inc in incidents:
            date_str = inc['date'].strftime('%d %b') if inc.get('date') else ''
            category = inc.get('category', '')
            matched = inc.get('matched', '')[:30]
            severity = inc.get('severity', '')

            # Severity indicator
            if severity == "high":
                sev_icon = "🔴"
                sev_color = "rgba(239,68,68,0.2)"
            elif severity == "medium":
                sev_icon = "🟡"
                sev_color = "rgba(245,158,11,0.2)"
            else:
                sev_icon = "🟢"
                sev_color = "rgba(34,197,94,0.2)"

            # Shorten category names
            cat_short = {
                'Physical Violence': 'Violence',
                'Verbal Aggression': 'Verbal',
                'Self-Harm': 'Self-Harm',
            }.get(category, category[:10])

            btn = QPushButton(f"{sev_icon} {date_str} [{cat_short}]: {matched}")
            btn.setCursor(Qt.PointingHandCursor)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {sev_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px 10px;
                    text-align: left;
                    color: #DDD;
                    font-size: 12px;
                }}
                QPushButton:hover {{
                    background: rgba(70,130,180,0.5);
                }}
            """)
            btn.clicked.connect(make_item_click_handler(inc))
            list_layout.addWidget(btn)

        list_layout.addStretch()
        scroll.setWidget(list_widget)
        layout.addWidget(scroll)

        # Size the popup
        popup.setFixedWidth(380)
        item_count = len(incidents)
        popup_height = min(50 + item_count * 34, 400)  # Max 400px height
        popup.setFixedHeight(popup_height)

        # Position near click
        global_pos = event.globalPosition().toPoint()
        popup.move(global_pos.x() - 190, global_pos.y() - 20)
        popup.show()

        # Wait for popup to close and handle selection
        popup.destroyed.connect(lambda: self._handle_popup_selection(popup._selected_data))

    def _handle_popup_selection(self, data):
        """Handle selection from incidents popup."""
        if data and data.get('date'):
            self.eventSelected.emit(data)


# ============================================================
# PROGRESS PANEL - Main floating panel with draggable narrative
# ============================================================

class ProgressPanel(QWidget):
    """Floating panel displaying progress timeline and narrative."""

    closed = Signal()

    def __init__(self, notes: List[Dict], parent=None, notes_panel=None, embedded=False):
        super().__init__(parent)
        self.notes = notes
        self.notes_panel = notes_panel
        self.embedded = embedded

        self._drag_offset = QPoint()
        self._dragging = False

        # Window settings - only for floating mode
        if not embedded:
            self.setWindowFlags(
                Qt.FramelessWindowHint |
                Qt.SubWindow |
                Qt.WindowStaysOnTopHint
            )
            self.setCursor(Qt.CursorShape.OpenHandCursor)
            self.resize(1000, 700)
            self.setMinimumSize(700, 500)
        else:
            # Allow shrinking when embedded
            self.setMinimumSize(1, 1)
            self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self._build_ui()
        self._analyze_and_display()

        try:
            apply_macos_blur(self)
        except:
            pass

        self.raise_()
        self.activateWindow()
        self.show()

    def _drag_start(self, e):
        if e.button() == Qt.LeftButton:
            self._drag_offset = e.globalPosition().toPoint() - self.frameGeometry().topLeft()

    def _drag_move(self, e):
        if e.buttons() & Qt.LeftButton:
            self.move(e.globalPosition().toPoint() - self._drag_offset)

    def mousePressEvent(self, evt):
        if self.embedded:
            return super().mousePressEvent(evt)
        if evt.button() == Qt.LeftButton:
            self._drag_offset = evt.globalPosition().toPoint() - self.frameGeometry().topLeft()
            self._dragging = True
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
        super().mousePressEvent(evt)

    def mouseMoveEvent(self, evt):
        if self.embedded:
            return super().mouseMoveEvent(evt)
        if getattr(self, '_dragging', False) and evt.buttons() & Qt.LeftButton:
            self.move(evt.globalPosition().toPoint() - self._drag_offset)
        super().mouseMoveEvent(evt)

    def mouseReleaseEvent(self, evt):
        if self.embedded:
            return super().mouseReleaseEvent(evt)
        if evt.button() == Qt.LeftButton:
            self._dragging = False
            self.setCursor(Qt.CursorShape.OpenHandCursor)
        super().mouseReleaseEvent(evt)

    def _build_ui(self):
        # Different styles for embedded vs floating
        if self.embedded:
            self.setStyleSheet("""
                QWidget {
                    background-color: white;
                    color: #333;
                }
                QLabel { background: transparent; color: #333; }
                QPushButton {
                    background-color: rgba(0,0,0,0.08);
                    color: #333;
                    border-radius: 6px;
                }
                QPushButton:hover { background-color: rgba(0,0,0,0.15); }
            """)
        else:
            self.setStyleSheet("""
                QWidget {
                    background-color: rgba(32,32,32,0.25);
                    color: #DCE6FF;
                    border-radius: 12px;
                }
                QLabel { color: #DCE6FF; }
                QPushButton {
                    background-color: rgba(255,255,255,0.22);
                    color: white;
                    border-radius: 6px;
                }
                QPushButton:hover { background-color: rgba(255,255,255,0.35); }
            """)

        outer = QVBoxLayout(self)
        outer.setContentsMargins(12, 32, 12, 12)  # Increased top margin
        outer.setSpacing(12)

        self.bg = QWidget()
        if self.embedded:
            self.bg.setStyleSheet("background-color: white; border-radius: 0;")
        else:
            self.bg.setStyleSheet("background-color: rgba(20,20,20,0.18); border-radius: 12px;")
        bg_layout = QVBoxLayout(self.bg)
        bg_layout.setContentsMargins(0, 0, 0, 0)
        bg_layout.setSpacing(0)
        outer.addWidget(self.bg)

        # Title bar
        self.title_bar = QWidget()
        self.title_bar.setFixedHeight(46)
        if not self.embedded:
            self.title_bar.setCursor(Qt.CursorShape.OpenHandCursor)
            self.title_bar.setStyleSheet("""
                background-color: rgba(30,30,30,0.35);
                border-top-left-radius: 12px;
                border-top-right-radius: 12px;
            """)
        else:
            self.title_bar.setStyleSheet("""
                background-color: rgba(240,242,245,0.95);
                border-bottom: 1px solid #d0d5da;
            """)

        tb = QHBoxLayout(self.title_bar)
        tb.setContentsMargins(12, 4, 12, 4)

        title = QLabel("📈 Progress Report")
        if self.embedded:
            title.setStyleSheet("font-size: 18px; font-weight: bold; color: #333; background: transparent;")
        else:
            title.setStyleSheet("font-size: 20px; font-weight: bold; color: #F5F5F5;")
        tb.addWidget(title)

        self.summary_label = QLabel("")
        if self.embedded:
            self.summary_label.setStyleSheet("font-size: 13px; color: #666; background: transparent;")
        else:
            self.summary_label.setStyleSheet("font-size: 14px; color: #AAA;")
        tb.addWidget(self.summary_label)

        tb.addStretch()

        # Narrative filter dropdown
        filter_label = QLabel("Narrative:")
        if self.embedded:
            filter_label.setStyleSheet("font-size: 12px; color: #666; background: transparent;")
        else:
            filter_label.setStyleSheet("font-size: 12px; color: #AAA;")
        tb.addWidget(filter_label)

        self.narrative_filter = QComboBox()
        self.narrative_filter.addItem("Last 1 Year", "1_year")
        self.narrative_filter.addItem("Full Notes", "all")
        self.narrative_filter.setFixedWidth(110)
        self.narrative_filter.setFixedHeight(28)
        if self.embedded:
            self.narrative_filter.setStyleSheet("""
                QComboBox {
                    background-color: #f0f2f5;
                    color: #333; font-size: 12px;
                    padding: 4px 8px;
                    border: 1px solid #d0d5da;
                    border-radius: 6px;
                }
                QComboBox:hover { background-color: #e5e7eb; }
                QComboBox::drop-down { border: none; }
                QComboBox QAbstractItemView { background: white; color: #333; }
            """)
        else:
            self.narrative_filter.setStyleSheet("""
                QComboBox {
                    background-color: rgba(255,255,255,0.18);
                    color: #FFFFFF; font-size: 12px;
                    padding: 4px 8px;
                    border: 1px solid rgba(255,255,255,0.25);
                    border-radius: 6px;
                }
                QComboBox:hover { background-color: rgba(255,255,255,0.28); }
                QComboBox::drop-down { border: none; }
                QComboBox QAbstractItemView { background: #333; color: white; }
            """)
        self.narrative_filter.currentIndexChanged.connect(self._on_filter_changed)
        tb.addWidget(self.narrative_filter)

        tb.addSpacing(10)

        export_btn = QPushButton("Export to Word")
        export_btn.setFixedHeight(28)
        export_btn.clicked.connect(self._export_to_word)
        export_btn.setStyleSheet("""
            QPushButton {
                background-color: rgba(70,130,180,0.6);
                color: #FFFFFF; font-size: 13px;
                padding: 4px 12px;
                border: 1px solid rgba(255,255,255,0.25);
                border-radius: 6px;
            }
            QPushButton:hover { background-color: rgba(70,130,180,0.8); }
        """)
        tb.addWidget(export_btn)

        # Only add close button for floating mode
        if not self.embedded:
            close_btn = QPushButton("✕")
            close_btn.setFixedSize(34, 28)
            close_btn.clicked.connect(self.close)
            close_btn.setStyleSheet("""
                QPushButton {
                    background-color: rgba(255,255,255,0.18);
                    color: #FFFFFF; font-size: 18px; font-weight: bold;
                    border: 1px solid rgba(255,255,255,0.25);
                    border-radius: 6px;
                }
                QPushButton:hover { background-color: rgba(255,255,255,0.32); }
            """)
            tb.addWidget(close_btn)

            self.title_bar.mousePressEvent = self._drag_start
            self.title_bar.mouseMoveEvent = self._drag_move

        bg_layout.addWidget(self.title_bar)

        # Main content area with vertical splitter (timeline above, narrative below)
        self.content_splitter = QSplitter(Qt.Vertical)
        self.content_splitter.setHandleWidth(8)
        self.content_splitter.setStyleSheet("""
            QSplitter::handle {
                background: rgba(100,150,200,0.5);
                border-radius: 4px;
            }
            QSplitter::handle:hover {
                background: rgba(100,150,200,0.8);
            }
        """)
        bg_layout.addWidget(self.content_splitter)

        # Top section - Timeline and controls
        top_panel = QWidget()
        top_panel.setStyleSheet("background: transparent;")
        top_layout = QVBoxLayout(top_panel)
        top_layout.setContentsMargins(12, 12, 12, 6)
        top_layout.setSpacing(12)

        # Zoom buttons row
        zoom_row = QHBoxLayout()
        zoom_label = QLabel("Time Range:")
        if self.embedded:
            zoom_label.setStyleSheet("font-size: 13px; color: #666; background: transparent;")
        else:
            zoom_label.setStyleSheet("font-size: 13px; color: #AAA;")
        zoom_row.addWidget(zoom_label)

        self.zoom_buttons = []
        self.zoom_levels = [("3M", 0.25), ("6M", 0.5), ("1Y", 1), ("2Y", 2), ("5Y", 5), ("All", None)]
        for label, years in self.zoom_levels:
            btn = QPushButton(label)
            btn.setFixedSize(45, 26)
            btn.setCursor(Qt.PointingHandCursor)
            if self.embedded:
                btn.setStyleSheet("""
                    QPushButton {
                        background-color: rgba(0,0,0,0.08);
                        color: #333; font-size: 11px;
                        border-radius: 4px;
                        border: 1px solid rgba(0,0,0,0.15);
                    }
                    QPushButton:hover { background-color: rgba(70,130,180,0.3); }
                    QPushButton:checked { background-color: rgba(70,130,180,0.5); color: white; }
                """)
            else:
                btn.setStyleSheet("""
                    QPushButton {
                        background-color: rgba(255,255,255,0.15);
                        color: white; font-size: 11px;
                        border-radius: 4px;
                        border: 1px solid rgba(255,255,255,0.2);
                    }
                    QPushButton:hover { background-color: rgba(70,130,180,0.5); }
                    QPushButton:checked { background-color: rgba(70,130,180,0.7); }
                """)
            btn.setCheckable(True)
            if label == "All":
                btn.setChecked(True)
            btn.clicked.connect(lambda checked, y=years: self._update_timeline_zoom(y))
            zoom_row.addWidget(btn)
            self.zoom_buttons.append(btn)

        zoom_row.addStretch()
        top_layout.addLayout(zoom_row)

        # Timeline scroll area
        timeline_scroll = QScrollArea()
        timeline_scroll.setWidgetResizable(True)
        timeline_scroll.setStyleSheet("""
            QScrollArea { background-color: rgba(0,0,0,0); border: none; }
            QScrollBar:horizontal {
                background: rgba(255,255,255,0.10); height: 10px;
                border-radius: 5px;
            }
            QScrollBar::handle:horizontal {
                background: rgba(255,255,255,0.35);
                border-radius: 5px; min-width: 40px;
            }
            QScrollBar:vertical {
                background: rgba(255,255,255,0.10); width: 10px;
                border-radius: 5px;
            }
            QScrollBar::handle:vertical {
                background: rgba(255,255,255,0.35);
                border-radius: 5px; min-height: 40px;
            }
        """)
        timeline_scroll.viewport().setStyleSheet("background: transparent;")

        # Timeline chart widget
        self.timeline_chart = InteractiveProgressTimeline()
        self.timeline_chart.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self.timeline_chart.setStyleSheet("background: transparent;")
        self.timeline_chart.eventSelected.connect(self._scroll_to_note)
        timeline_scroll.setWidget(self.timeline_chart)

        top_layout.addWidget(timeline_scroll)
        self.content_splitter.addWidget(top_panel)

        # Bottom section - Narrative (resizable via vertical splitter drag bar)
        bottom_panel = QWidget()
        bottom_panel.setStyleSheet("background: transparent;")
        bottom_layout = QVBoxLayout(bottom_panel)
        bottom_layout.setContentsMargins(12, 6, 12, 12)
        bottom_layout.setSpacing(8)

        narrative_header = QLabel("📝 Narrative Summary")
        if self.embedded:
            narrative_header.setStyleSheet("font-size: 16px; font-weight: bold; color: #333; background: transparent;")
        else:
            narrative_header.setStyleSheet("font-size: 16px; font-weight: bold; color: #F5F5F5;")
        bottom_layout.addWidget(narrative_header)

        self.narrative_text = QTextBrowser()
        self.narrative_text.setReadOnly(True)
        self.narrative_text.setOpenLinks(False)  # Handle clicks manually
        self.narrative_text.anchorClicked.connect(self._on_narrative_link_clicked)
        if self.embedded:
            self.narrative_text.setStyleSheet("""
                QTextBrowser {
                    background-color: #f8f9fa;
                    border: 1px solid #e0e0e0;
                    border-radius: 8px; padding: 12px;
                    color: #333; font-size: 13px;
                    font-family: monospace;
                }
            """)
        else:
            self.narrative_text.setStyleSheet("""
                QTextBrowser {
                    background-color: rgba(40,40,40,0.5);
                    border-radius: 8px; padding: 12px;
                    color: #DCE6FF; font-size: 13px;
                    font-family: monospace;
                }
            """)
        bottom_layout.addWidget(self.narrative_text)

        self.content_splitter.addWidget(bottom_panel)

        # Set initial splitter sizes (timeline takes top portion, narrative below)
        self.content_splitter.setSizes([400, 250])

        # Resize grip - only for floating mode
        if not self.embedded:
            self.resize_grip = QSizeGrip(self)
            self.resize_grip.setFixedSize(20, 20)
            self.resize_grip.setStyleSheet("""
                QSizeGrip {
                    background-color: rgba(100,150,200,0.6);
                    border-radius: 4px;
                    border: 1px solid rgba(255,255,255,0.3);
                }
                QSizeGrip:hover {
                    background-color: rgba(100,150,200,0.9);
                }
            """)
            self.resize_grip.setCursor(Qt.SizeFDiagCursor)
            bg_layout.addWidget(self.resize_grip, alignment=Qt.AlignBottom | Qt.AlignRight)
        else:
            self.resize_grip = None

    def _on_filter_changed(self, index):
        """Handle narrative filter dropdown change."""
        self._regenerate_narrative()

    def _analyze_and_display(self):
        self.results = analyze_notes_for_progress(self.notes)

        total = self.results["total_notes"]
        events = len(self.results["tentpole_events"])
        self.summary_label.setText(f"• {total} notes analyzed • {events} key events")

        # Generate timeline (all data initially)
        self._update_timeline_zoom(None)

        # Generate narrative with current filter
        self._regenerate_narrative()

    def _regenerate_narrative(self):
        """Regenerate narrative based on current filter selection."""
        # Get current filter period from dropdown
        period = self.narrative_filter.currentData() or '1_year'

        # Get patient info including gender from shared store
        store = get_shared_store()
        patient_info = store.patient_info
        patient_name = patient_info.get("name", "The patient")

        # Get gender - convert to 'M', 'F', or None
        gender_raw = patient_info.get("gender", "").lower()
        if gender_raw in ("male", "m"):
            gender = "M"
        elif gender_raw in ("female", "f"):
            gender = "F"
        else:
            gender = None

        # Generate narrative using narrative_generator module with filtering
        try:
            from narrative_generator import generate_narrative as gen_narrative, filter_entries_by_period, get_date_range_info

            # Convert notes to entry format
            entries = []
            for note in self.notes:
                date = note.get('date') or note.get('datetime')
                content = note.get('content', note.get('text', ''))
                if date and content:
                    entries.append({
                        'date': date,
                        'content': content,
                        'text': content,
                        'type': note.get('type', ''),
                        'originator': note.get('originator', ''),
                    })

            if entries:
                # Generate narrative with period filter
                plain_text, html_narrative = gen_narrative(entries, period=period)

                # Get date range info for display
                date_range = get_date_range_info(entries, period=period)
                # Update summary label with date range
                total = self.results["total_notes"]
                events = len(self.results["tentpole_events"])
                filtered_count = len(filter_entries_by_period(entries, period))
                self.summary_label.setText(f"• {filtered_count}/{total} notes • {events} events • {date_range}")
            else:
                # Fallback to original if no entries
                narrative_text = generate_narrative(self.results, patient_name=patient_name, gender=gender)
                html_narrative = narrative_text.replace("\n", "<br>")

        except Exception as e:
            print(f"[ProgressPanel] Failed to use narrative_generator, falling back: {e}")
            import traceback
            traceback.print_exc()
            # Fallback to original narrative generator
            narrative_text = generate_narrative(self.results, patient_name=patient_name, gender=gender)
            html_narrative = narrative_text.replace("\n", "<br>")
            html_narrative = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', html_narrative)

        # Add disclaimer at the top
        disclaimer = "<i>The following account is a guide - please check the narrative presented against the notes.</i><br><br>"
        html_narrative = disclaimer + html_narrative

        # Use dark text for embedded mode, light for floating
        text_color = "#333" if self.embedded else "#DCE6FF"
        link_color = "#0066cc" if self.embedded else "#66b3ff"
        link_hover_bg = "rgba(255,200,0,0.3)" if self.embedded else "rgba(255,200,0,0.4)"
        self.narrative_text.setHtml(f"""
            <style>
                a {{ color: {link_color}; text-decoration: underline; }}
                a:hover {{ background-color: {link_hover_bg}; }}
            </style>
            <div style='font-family: monospace; font-size: 13px; color: {text_color};'>{html_narrative}</div>
        """)

    def _scroll_to_note(self, event_data: dict):
        """Scroll to the note in the left panel and highlight relevant text."""
        if not self.notes_panel:
            return

        date = event_data.get("date")
        matched = event_data.get("matched", "")
        content_snippet = event_data.get("content_snippet", "")

        if date:
            try:
                # Use jump_to_note with content snippet for precise navigation
                if content_snippet and hasattr(self.notes_panel, 'jump_to_note'):
                    self.notes_panel.jump_to_note(date, content_snippet)
                elif hasattr(self.notes_panel, 'jump_to_date'):
                    self.notes_panel.jump_to_date(date)

                # Highlight the matched text if available
                if matched and hasattr(self.notes_panel, 'highlight_text'):
                    self.notes_panel.highlight_text(matched)
            except Exception as e:
                print(f"Error scrolling to note: {e}")

    def _on_narrative_link_clicked(self, url: QUrl):
        """Handle clicks on narrative reference links."""
        ref_id = url.fragment()  # Gets the part after #
        if not ref_id:
            return

        ref_data = get_reference(ref_id)
        if not ref_data:
            return

        # Reuse existing _scroll_to_note logic
        self._scroll_to_note(ref_data)

    def _update_timeline_zoom(self, years):
        """Update timeline to show specified time range."""
        # Update button checked states
        year_map = {0.25: 0, 0.5: 1, 1: 2, 2: 3, 5: 4, None: 5}
        btn_index = year_map.get(years, 5)
        for i, btn in enumerate(self.zoom_buttons):
            btn.setChecked(i == btn_index)

        all_months = self.results["all_months"]
        if not all_months:
            return

        # Filter data based on years
        if years is None:
            filtered_months = all_months
        else:
            cutoff = datetime.now() - timedelta(days=years * 365)
            cutoff_key = cutoff.strftime("%Y-%m")
            filtered_months = [m for m in all_months if m >= cutoff_key]

        if not filtered_months or len(filtered_months) < 2:
            self.timeline_chart.setText(f"Insufficient data for timeline")
            return

        # Create filtered results with ALL required keys
        filtered_results = {
            "all_months": filtered_months,
            "monthly_data": {m: self.results["monthly_data"][m] for m in filtered_months},
            "tentpole_events": [e for e in self.results["tentpole_events"] if e["month"] in filtered_months],
            "monthly_violence": {m: c for m, c in self.results.get("monthly_violence", {}).items() if m in filtered_months},
            "monthly_verbal": {m: c for m, c in self.results.get("monthly_verbal", {}).items() if m in filtered_months},
            "monthly_incidents": {m: c for m, c in self.results.get("monthly_incidents", {}).items() if m in filtered_months},
        }

        chart_bytes, timeline_info = create_progress_timeline_chart(filtered_results)
        if chart_bytes:
            # Pass monthly_data for incident popup
            self.timeline_chart.set_chart(chart_bytes, timeline_info, filtered_results.get("monthly_data", {}))
        else:
            self.timeline_chart.setText("Insufficient data for timeline")

    def _export_to_word(self):
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        import os
        import re

        default_name = f"Progress_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Progress Report", default_name, "Word Document (*.docx)"
        )

        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading('Progress Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Get patient info including gender from shared store
            store = get_shared_store()
            patient_info = store.patient_info
            patient_name = patient_info.get("name", "The patient")

            # Get gender - convert to 'M', 'F', or None
            gender_raw = patient_info.get("gender", "").lower()
            if gender_raw in ("male", "m"):
                gender = "M"
            elif gender_raw in ("female", "f"):
                gender = "F"
            else:
                gender = None

            # Get narrative and parse into sections
            narrative = generate_narrative(self.results, patient_name=patient_name, gender=gender)

            # Add disclaimer at the top
            disclaimer_para = doc.add_paragraph()
            disclaimer_run = disclaimer_para.add_run("The following account is a guide - please check the narrative presented against the notes.")
            disclaimer_run.italic = True
            doc.add_paragraph()  # Empty line after disclaimer

            def clean_line(text):
                """Strip HTML tags and convert HTML entities."""
                # Remove anchor tags but keep inner text
                text = re.sub(r'<a[^>]*>([^<]*)</a>', r'\1', text)
                # Remove any other HTML tags
                text = re.sub(r'<[^>]+>', '', text)
                # Convert HTML entities
                text = text.replace('&amp;', '&')
                text = text.replace('&lt;', '<')
                text = text.replace('&gt;', '>')
                text = text.replace('&quot;', '"')
                return text

            def add_line_with_bold(paragraph, text):
                """Add text to paragraph, converting **text** to bold runs."""
                # Split by bold markers
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    elif part:
                        paragraph.add_run(part)

            # Parse narrative into sections
            lines = narrative.split('\n')
            for line in lines:
                if line.startswith('='):
                    continue
                elif line.startswith('**') and line.endswith('**'):
                    # Date header - make it bold
                    heading_text = line.strip('*')
                    p = doc.add_paragraph()
                    run = p.add_run(heading_text)
                    run.bold = True
                    run.font.size = Pt(12)
                elif line.strip():
                    cleaned = clean_line(line)
                    p = doc.add_paragraph()
                    add_line_with_bold(p, cleaned)

            doc.save(file_path)

            QMessageBox.information(self, "Export Complete", f"Report saved to:\n{file_path}")
            os.system(f'open "{file_path}"')

        except Exception as e:
            QMessageBox.warning(self, "Export Failed", f"Could not export: {e}")
