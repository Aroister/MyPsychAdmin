"""
GPR Report Parser — Parse medical/psychiatric tribunal report DOCX files for GPR population.

Supports three DOCX formats:
  1. table_form   — T131 RC Report style (questions in table cells, answers adjacent)
  2. numbered_paragraphs — JB tribunal style (numbered sections as paragraphs)
  3. narrative    — SUSSEX/SWLSTG style (prose headings, content paragraphs)
"""

import os
import re
from datetime import datetime


# ============================================================
# T131 Section Number → GPR Card Key
# ============================================================
T131_TO_GPR = {
    1: "patient_details",
    2: "signature",
    3: "legal_criteria",       # Factors affecting hearing
    4: "legal_criteria",       # Adjustments
    5: "forensic",             # Index offence & forensic history
    6: "psych_history",        # Previous MH dates
    7: "psych_history",        # Previous admission reasons
    8: "circumstances",        # Current admission circumstances
    9: "diagnosis",            # Mental disorder/diagnosis
    10: "diagnosis",           # Learning disability
    11: "legal_criteria",      # Detention required
    12: "medication",          # Treatment/medication
    13: "strengths",           # Strengths/positive factors
    14: "circumstances",       # Current progress
    15: "medication",          # Compliance
    16: "legal_criteria",      # MCA/DoL
    17: "risk",                # Harm to self/others
    18: "risk",                # Property damage
    19: "legal_criteria",      # S2 detention justified
    20: "legal_criteria",      # Other detention justified
    21: "risk",                # Discharge risk
    22: "legal_criteria",      # Community risk management
    23: "legal_criteria",      # Recommendations
    24: "signature",           # Signature
}

# Sub-header labels for sections that merge multiple T131 items
T131_SUB_HEADERS = {
    3: "Factors Affecting Hearing",
    4: "Adjustments",
    5: "Index Offence & Forensic History",
    6: "Previous Mental Health Dates",
    7: "Previous Admission Reasons",
    8: "Current Admission Circumstances",
    9: "Mental Disorder / Diagnosis",
    10: "Learning Disability",
    11: "Detention Required",
    12: "Treatment / Medication",
    13: "Strengths / Positive Factors",
    14: "Current Progress",
    15: "Compliance with Treatment",
    16: "Mental Capacity Act / DoL",
    17: "Incidents of Harm to Self/Others",
    18: "Incidents of Property Damage",
    19: "Section 2 Detention",
    20: "Other Detention",
    21: "Risk if Discharged",
    22: "Community Risk Management",
    23: "Recommendations",
}

# Heading patterns → T131 card key (used for all formats)
HEADING_PATTERNS = [
    # Factors affecting hearing (3)
    (r'factors.*affect.*understanding', 'factors_hearing'),
    (r'ability.*cope.*hearing', 'factors_hearing'),

    # Adjustments (4)
    (r'adjustments.*(?:panel|tribunal).*consider', 'adjustments'),

    # Index offence / forensic (5)
    (r'index offence', 'forensic'),
    (r'forensic.*history', 'forensic'),
    (r'relevant.*forensic', 'forensic'),

    # Previous MH dates (6)
    (r'dates.*previous.*mental health', 'previous_mh_dates'),
    (r'previous.*involvement.*mental health', 'previous_mh_dates'),

    # Previous admission reasons (7)
    (r'reasons.*previous.*admission', 'previous_admission_reasons'),
    (r'give reasons.*previous', 'previous_admission_reasons'),

    # Current admission (8)
    (r'circumstances.*current.*admission', 'current_admission'),
    (r'current admission', 'current_admission'),

    # Diagnosis (9) — must come before learning disability
    (r'is the patient.*suffering', 'diagnosis'),
    (r'mental disorder.*nature', 'diagnosis'),
    (r'9\.\s*mental disorder', 'diagnosis'),
    (r'nature.*degree.*mental disorder', 'diagnosis'),
    (r'from.*mental disorder', 'diagnosis'),

    # Learning disability (10)
    (r'learning disability', 'learning_disability'),
    (r'abnormally aggressive', 'learning_disability'),

    # Detention requirement (11)
    (r'what is it.*necessary', 'detention_required'),
    (r'necessary.*medical treatment', 'detention_required'),
    (r'why is detention necessary', 'detention_required'),

    # Treatment / medication (12)
    (r'medical treatment.*prescribed', 'treatment'),
    (r'appropriate.*available.*treatment', 'treatment'),
    (r'what.*treatment', 'treatment'),

    # Strengths (13)
    (r'strengths.*positive factors', 'strengths'),
    (r'what are the strengths', 'strengths'),

    # Progress (14)
    (r'current progress', 'progress'),
    (r'summary.*progress', 'progress'),
    (r'progress.*behaviour', 'progress'),

    # Compliance (15)
    (r'understanding.*compliance', 'compliance'),
    (r'compliance.*willingness', 'compliance'),
    (r'willingness to accept', 'compliance'),
    (r"patient's understanding of", 'compliance'),
    (r'prescribed medication.*mental disorder', 'compliance'),

    # MCA/DoL (16)
    (r'mental capacity act', 'mca_dol'),
    (r'deprivation of liberty', 'mca_dol'),

    # Risk harm (17)
    (r'harmed themselves or others', 'risk_harm'),
    (r'incidents.*harm', 'risk_harm'),
    (r'threatened to harm', 'risk_harm'),

    # Risk property (18)
    (r'damaged property', 'risk_property'),
    (r'threatened to damage property', 'risk_property'),

    # S2 detention (19)
    (r'section 2 cases.*detention', 's2_detention'),
    (r'in section 2 cases', 's2_detention'),

    # Other detention (20)
    (r'all other cases.*provision', 'other_detention'),
    (r'in all other cases', 'other_detention'),

    # Discharge risk (21)
    (r'discharged.*dangerous', 'discharge_risk'),
    (r'likely to act in a manner dangerous', 'discharge_risk'),

    # Community (22)
    (r'risks.*managed.*community', 'community'),
    (r'managed effectively in the community', 'community'),

    # Recommendations (23)
    (r'recommendations.*tribunal', 'recommendations'),
    (r'do you have any recommendations', 'recommendations'),
]

# T131 intermediate key → GPR card key
_INTERMEDIATE_TO_GPR = {
    'factors_hearing': 'legal_criteria',
    'adjustments': 'legal_criteria',
    'forensic': 'forensic',
    'previous_mh_dates': 'psych_history',
    'previous_admission_reasons': 'psych_history',
    'current_admission': 'circumstances',
    'diagnosis': 'diagnosis',
    'learning_disability': 'diagnosis',
    'detention_required': 'legal_criteria',
    'treatment': 'medication',
    'strengths': 'strengths',
    'progress': 'circumstances',
    'compliance': 'medication',
    'mca_dol': 'legal_criteria',
    'risk_harm': 'risk',
    'risk_property': 'risk',
    's2_detention': 'legal_criteria',
    'other_detention': 'legal_criteria',
    'discharge_risk': 'risk',
    'community': 'legal_criteria',
    'recommendations': 'legal_criteria',
}

# Intermediate key → sub-header label for merged sections
_INTERMEDIATE_LABELS = {
    'factors_hearing': 'Factors Affecting Hearing',
    'adjustments': 'Adjustments',
    'forensic': 'Index Offence & Forensic History',
    'previous_mh_dates': 'Previous Mental Health Dates',
    'previous_admission_reasons': 'Previous Admission Reasons',
    'current_admission': 'Current Admission Circumstances',
    'diagnosis': 'Mental Disorder / Diagnosis',
    'learning_disability': 'Learning Disability',
    'detention_required': 'Detention Required',
    'treatment': 'Treatment / Medication',
    'strengths': 'Strengths / Positive Factors',
    'progress': 'Current Progress',
    'compliance': 'Compliance with Treatment',
    'mca_dol': 'Mental Capacity Act / DoL',
    'risk_harm': 'Incidents of Harm to Self/Others',
    'risk_property': 'Incidents of Property Damage',
    's2_detention': 'Section 2 Detention',
    'other_detention': 'Other Detention',
    'discharge_risk': 'Risk if Discharged',
    'community': 'Community Risk Management',
    'recommendations': 'Recommendations',
}

# Narrative heading keywords → GPR card key
NARRATIVE_HEADING_MAP = [
    # Must be ordered: specific before generic
    (r'sources?\s+of\s+information', 'report_based_on'),
    (r'informant|based\s+on', 'report_based_on'),
    (r'knowledge\s+of\s+patient', 'report_based_on'),
    (r'psychiatric\s+history.*(?:history|admissions)', 'psych_history'),
    (r'past\s+psychiatric', 'psych_history'),
    (r'psychiatric\s+history', 'psych_history'),
    (r'mental\s+health\s+history', 'psych_history'),
    (r'medical\s+history|physical\s+health', 'medical_history'),
    (r'past\s+medical', 'medical_history'),
    (r'forensic\s+history', 'forensic'),
    (r'offending|index\s+offence', 'forensic'),
    (r'substance\s+(?:use|misuse|abuse)|alcohol|drug', 'substance_use'),
    (r'medication|treatment|pharmacological', 'medication'),
    (r'risk\s+assessment|risk\s+management|risk\s+factors', 'risk'),
    (r'diagnosis|formulation|mental\s+state\s+exam', 'diagnosis'),
    (r'strengths|positive\s+factors|protective', 'strengths'),
    (r'statutory\s+criteria', 'legal_criteria'),
    (r'legal\s+criteria|detention\s+criteria', 'legal_criteria'),
    (r'summary\s+of\s+case|case\s+for\s+maintaining', 'legal_criteria'),
    (r'relevance\s+of\s+mca|mental\s+capacity', 'legal_criteria'),
    (r'capacity', 'legal_criteria'),
    (r'background|personal\s+history|family\s+history|social\s+history|early\s+life', 'background'),
    (r'presenting\s+complaint|admission|current\s+episode|circumstances', 'circumstances'),
    (r'progress|current\s+mental\s+state|mental\s+state', 'circumstances'),
    (r'recommendation|opinion|conclusion', 'legal_criteria'),
    (r'signature|signed|declaration', 'signature'),
]

# Question patterns to exclude from content
QUESTION_PATTERNS = [
    'are there any factors', 'are there any adjustments',
    'what is the nature of', 'give details of any',
    'what are the strengths', 'give a summary of',
    'in section 2 cases', 'in all other cases',
    'if the patient was discharged', 'if the patient were discharged',
    'please explain how', 'is there any other relevant',
    'do you have any recommendations', 'is the patient now suffering',
    'what appropriate and available', 'what are the dates',
    'what are the circumstances', 'give reasons for',
    'does the patient have a learning', 'what is it about',
    'would they be likely to act', 'managed effectively in the community',
    'if yes, has a diagnosis', 'if yes, what is the diagnosis',
    'has a diagnosis been made', 'what is the diagnosis',
]


# ============================================================
# Content Cleaning Helpers
# ============================================================

def _clean_content(text: str) -> str:
    """Remove checkbox symbols, normalize whitespace."""
    if not text:
        return ""
    text = re.sub(r'\[\s*[xX]?\s*\]', '', text)
    text = re.sub(r'[☐☒☑✓✔]', '', text)
    text = re.sub(r'–\s*If yes[^:?]*[:\?]?\s*', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_checkbox_answer(text: str) -> str:
    """Extract content after [x]/checked boxes."""
    # [x] pattern
    match = re.search(
        r'\[[\s]*[xX][\s]*\][^a-zA-Z]*(?:If yes[^\n]*\n+)?(.+)',
        text, re.DOTALL | re.IGNORECASE
    )
    if match:
        content = match.group(1).strip()
        if content and len(content) > 5:
            return content
    # ☒ pattern
    match = re.search(
        r'☒[^a-zA-Z]*(?:If yes[^\n]*\n+)?(.+)',
        text, re.DOTALL | re.IGNORECASE
    )
    if match:
        content = match.group(1).strip()
        if content and len(content) > 5:
            return content
    return ""


def _is_question_text(text: str) -> bool:
    """Check if text is a question/heading rather than answer content."""
    if not text or len(text) < 10:
        return False
    lower = re.sub(r'^[\[\]☐☒xX\s\-–\d\.]*', '', text.lower().strip()).strip()
    for pattern in QUESTION_PATTERNS:
        if lower.startswith(pattern):
            return True
    if re.match(r'^\d{1,2}\.\s+', text):
        return True
    return False


def _is_checkbox_only(text: str) -> bool:
    """Check if text is just checkbox format without actual content."""
    if not text:
        return False
    cleaned = re.sub(r'[\[\]☐☒xX\s\n]', '', text)
    cleaned = re.sub(r'(?:No|Yes|N/A)', '', cleaned, flags=re.IGNORECASE)
    return len(cleaned.strip()) < 5


def _extract_patient_info(sections: dict) -> dict:
    """Parse patient details from parsed sections into a patient_info dict."""
    patient_info = {
        "name": None, "dob": None, "nhs_number": None,
        "address": None, "gender": None, "mha_status": None,
        "age": None, "ethnicity": None,
    }

    text = sections.get("patient_details", "")
    if not text:
        return None

    # Name
    m = re.search(r'(?:Name|PATIENT)[:\s]+([A-Za-z][A-Za-z\s\-\']+?)(?:\n|$)', text, re.IGNORECASE)
    if m:
        patient_info["name"] = m.group(1).strip()

    # DOB — handle ordinal suffixes and commas
    m = re.search(
        r'(?:DOB|Date\s*of\s*Birth)[:\s]+(.+?)(?:\n|$)',
        text, re.IGNORECASE
    )
    if m:
        dob_str = m.group(1).strip()
        # Remove ordinal suffixes
        dob_str = re.sub(r'(\d+)(?:st|nd|rd|th)', r'\1', dob_str)
        dob_str = dob_str.replace(',', '')
        for fmt in ('%d %B %Y', '%d/%m/%Y', '%d-%m-%Y', '%d %b %Y', '%Y-%m-%d'):
            try:
                patient_info["dob"] = datetime.strptime(dob_str.strip(), fmt)
                break
            except ValueError:
                continue

    # NHS number
    m = re.search(r'(?:NHS|NHS\s*Number)[:\s]+(\d[\d\s]+)', text, re.IGNORECASE)
    if m:
        patient_info["nhs_number"] = m.group(1).strip()

    # Address
    m = re.search(r'(?:Address|Residence)[:\s]+(.+?)(?:\n|$)', text, re.IGNORECASE)
    if m:
        patient_info["address"] = m.group(1).strip()

    # Gender
    m = re.search(r'(?:Gender|Sex)[:\s]+(Male|Female|Other)', text, re.IGNORECASE)
    if m:
        patient_info["gender"] = m.group(1).title()

    # MHA Status
    m = re.search(r'(?:MHA\s*Status|Section|Detained\s*Under)[:\s]+(.+?)(?:\n|$)', text, re.IGNORECASE)
    if m:
        patient_info["mha_status"] = m.group(1).strip()

    # Age — calculate from DOB if not explicit
    if patient_info["dob"] and not patient_info["age"]:
        today = datetime.today()
        dob = patient_info["dob"]
        age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
        if 0 < age < 120:
            patient_info["age"] = age

    if any(patient_info.values()):
        return patient_info
    return None


# ============================================================
# Format Detection
# ============================================================

def detect_format(doc) -> str:
    """Returns 'table_form', 'numbered_paragraphs', or 'narrative'."""
    num_tables = len(doc.tables)

    # Check for "RESPONSIBLE CLINICIAN" in first few paragraphs
    has_rc_marker = False
    for para in doc.paragraphs[:15]:
        if 'responsible clinician' in para.text.lower():
            has_rc_marker = True
            break

    # Count numbered paragraphs
    numbered_count = 0
    for para in doc.paragraphs:
        if re.match(r'^\d{1,2}\.\s', para.text.strip()):
            numbered_count += 1

    # table_form: RC marker + enough tables
    if has_rc_marker and num_tables >= 4:
        return 'table_form'

    # numbered_paragraphs: few tables + many numbered sections
    if num_tables < 3 and numbered_count >= 5:
        return 'numbered_paragraphs'

    # table_form: many tables (even without RC marker — some forms omit it)
    if num_tables >= 8:
        return 'table_form'

    # numbered_paragraphs with some tables
    if numbered_count >= 3 and num_tables < 8:
        return 'numbered_paragraphs'

    # narrative: everything else
    return 'narrative'


# ============================================================
# Table-Form Parser (Format 1)
# ============================================================

def _match_heading(text: str) -> str:
    """Match heading text to an intermediate T131 key."""
    if not text or len(text) < 10:
        return None
    lower = text.lower()
    for pattern, key in HEADING_PATTERNS:
        if re.search(pattern, lower):
            return key
    return None


def _get_unique_cells(cell_list):
    """Deduplicate adjacent cells (merged cells appear multiple times)."""
    unique = []
    for c in cell_list:
        if c and (not unique or c != unique[-1]):
            unique.append(c)
    return unique


def _match_patient_detail_row(lower0: str, val: str, patient_parts: list, author_parts: list):
    """Match a table row label to patient detail or author fields.

    Handles both formal labels ('Name of Patient') and simple ones ('Name').
    """
    if not val or not val.strip():
        return

    # Patient name — "Name", "Name of Patient", "Patient Name", "Patient information\nName"
    if (lower0 == 'name' or 'name of patient' in lower0 or 'patient name' in lower0
            or lower0.endswith('\nname') or lower0 == 'patient'):
        patient_parts.append(f"Name: {val}")
    elif 'date of birth' in lower0 or lower0 == 'dob':
        patient_parts.append(f"DOB: {val}")
    elif 'nhs' in lower0 or 'rio' in lower0:
        patient_parts.append(f"NHS: {val}")
    elif 'address' in lower0 or 'residence' in lower0:
        patient_parts.append(f"Address: {val}")
    elif 'mental health act' in lower0 or 'mha status' in lower0:
        patient_parts.append(f"MHA Status: {val}")
    elif 'current section' in lower0 or 'detained under' in lower0:
        patient_parts.append(f"MHA Status: {val}")
    elif lower0 == 'gender' or lower0 == 'sex':
        patient_parts.append(f"Gender: {val}")
    elif 'date of admission' in lower0:
        patient_parts.append(f"Date of Admission: {val}")
    elif 'current diagnosis' in lower0:
        patient_parts.append(f"Diagnosis: {val}")
    elif 'current medication' in lower0:
        patient_parts.append(f"Medication: {val}")
    elif 'registered gp' in lower0:
        patient_parts.append(f"GP: {val}")
    # Author / RC fields
    elif 'your name' in lower0 or 'name of rc' in lower0 or 'author of report' in lower0:
        author_parts.append(f"Name: {val}")
    elif 'your role' in lower0 or 'role' == lower0:
        author_parts.append(f"Role: {val}")
    elif 'date of report' in lower0:
        author_parts.append(f"Date: {val}")


def _parse_table_form(doc) -> dict:
    """Parse T131 table-form DOCX. Returns {intermediate_key: content}."""
    sections = {}
    patient_details_parts = []
    author_parts = []

    for table in doc.tables:
        rows = list(table.rows)
        i = 0
        while i < len(rows):
            cells = [cell.text.strip() for cell in rows[i].cells]
            if not any(cells):
                i += 1
                continue

            unique_cells = _get_unique_cells(cells)

            # Try heading match
            heading_key = None
            heading_cell_idx = -1
            for cell_idx, cell_text in enumerate(unique_cells):
                if re.match(r'^\d{1,2}\.\s*$', cell_text):
                    continue
                matched = _match_heading(cell_text)
                if matched:
                    heading_key = matched
                    heading_cell_idx = cell_idx
                    break

            if heading_key:
                answer = ""

                # Check cells after heading in same row
                start_idx = heading_cell_idx + 1 if heading_cell_idx >= 0 else 1
                if len(unique_cells) > start_idx:
                    for cell_text in unique_cells[start_idx:]:
                        if _is_question_text(cell_text):
                            continue
                        if cell_text and not re.match(r'^\d+\.\s*', cell_text):
                            cleaned = _clean_content(cell_text)
                            if cleaned and cleaned not in ('No', 'Yes', 'N/A'):
                                answer = cleaned
                                break

                # Check next row
                if not answer and i + 1 < len(rows):
                    next_cells = [cell.text.strip() for cell in rows[i + 1].cells]
                    unique_next = _get_unique_cells(next_cells)

                    for cell_text in unique_next:
                        if re.match(r'^\d+\.\s*', cell_text):
                            break
                        if _is_question_text(cell_text):
                            continue
                        if cell_text.strip() in ('No\nYes', 'Yes\nNo', 'No', 'Yes', 'N/A'):
                            continue

                        # Checkbox extraction
                        if '[x' in cell_text.lower() or '☒' in cell_text or '[' in cell_text:
                            yes_answer = _extract_checkbox_answer(cell_text)
                            if yes_answer:
                                answer = yes_answer
                                break

                        cleaned = _clean_content(cell_text)
                        if cleaned and not _is_question_text(cleaned):
                            if cleaned.replace('\n', ' ').strip() in ('No Yes', 'Yes No'):
                                continue
                            if _is_checkbox_only(cleaned):
                                continue
                            answer = cleaned
                            break

                    if answer:
                        i += 1

                if answer:
                    # Add Yes prefix for yes/no sections
                    yes_no_keys = {'factors_hearing', 'adjustments', 's2_detention',
                                   'other_detention', 'discharge_risk', 'recommendations'}
                    if heading_key in yes_no_keys and answer not in ('Yes', 'No', 'N/A'):
                        if not answer.startswith('Yes') and not answer.startswith('No'):
                            answer = f"Yes - {answer}"
                    sections[heading_key] = answer

            # Patient details fields — match simple labels too (e.g. "Name", not just "Name of Patient")
            else:
                lower0 = cells[0].lower().strip()
                if len(unique_cells) > 1:
                    val = unique_cells[1]
                    _match_patient_detail_row(lower0, val, patient_details_parts, author_parts)

            i += 1

    if patient_details_parts:
        sections['patient_details_raw'] = '\n'.join(patient_details_parts)
    if author_parts:
        sections['author_raw'] = '\n'.join(author_parts)

    return sections


# ============================================================
# Numbered-Paragraphs Parser (Format 2)
# ============================================================

def _parse_numbered_paragraphs(doc) -> dict:
    """Parse numbered-paragraphs style DOCX. Returns {intermediate_key: content}."""
    sections = {}
    current_key = None
    current_content = []
    patient_details_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect numbered section header
        section_match = re.match(r'^(\d{1,2})\.\s+(.*)', text)

        # Special detection for unnumbered section 2
        is_section_2 = (
            'factors' in text.lower() and 'affect' in text.lower()
            and 'understanding' in text.lower()
        )

        if section_match or is_section_2:
            # Save previous section
            if current_key and current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    sections[current_key] = content

            # Match heading to key
            if section_match:
                heading_text = section_match.group(2) if section_match.group(2) else text
            else:
                heading_text = text

            current_key = _match_heading(heading_text)
            if not current_key and section_match:
                # Fallback: use T131 section number
                sec_num = int(section_match.group(1))
                gpr_key = T131_TO_GPR.get(sec_num)
                if gpr_key:
                    # Use the section number as intermediate key
                    current_key = f"sec_{sec_num}"
            current_content = []
            continue

        # Detect checkbox answers
        if current_key:
            # Check for checked checkbox
            if text.startswith('☒') or re.match(r'\[\s*[xX]\s*\]', text):
                # Extract content after Yes marker
                yes_match = re.match(
                    r'(?:☒|\[\s*[xX]\s*\])\s*Yes\s*[-–:]\s*(.+)',
                    text, re.IGNORECASE
                )
                if yes_match:
                    current_content.append(yes_match.group(1).strip())
                elif 'yes' in text.lower():
                    current_content.append("Yes")
                else:
                    cleaned = _clean_content(text)
                    if cleaned:
                        current_content.append(cleaned)
            elif text.startswith('☐') or re.match(r'\[\s*\]', text):
                # Unchecked — skip
                pass
            else:
                # Regular content line
                cleaned = _clean_content(text)
                if cleaned and not _is_question_text(cleaned):
                    current_content.append(cleaned)

    # Save last section
    if current_key and current_content:
        content = '\n'.join(current_content).strip()
        if content:
            sections[current_key] = content

    # Also parse tables for patient details
    author_parts = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells:
                continue
            lower0 = cells[0].lower().strip()
            unique = _get_unique_cells(cells)
            if len(unique) > 1:
                _match_patient_detail_row(lower0, unique[1], patient_details_parts, author_parts)

    if patient_details_parts:
        sections['patient_details_raw'] = '\n'.join(patient_details_parts)
    if author_parts:
        sections['author_raw'] = '\n'.join(author_parts)

    return sections


# ============================================================
# Narrative Parser (Format 3)
# ============================================================

def _parse_narrative(doc) -> dict:
    """Parse narrative-style DOCX with prose headings. Returns {gpr_key: content}."""
    sections = {}
    patient_details_parts = []

    # Extract patient details from first few tables
    author_parts = []
    for table in doc.tables[:5]:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2:
                continue
            unique = _get_unique_cells(cells)
            if len(unique) < 2:
                continue
            lower0 = cells[0].lower().strip()
            val = unique[1]
            _match_patient_detail_row(lower0, val, patient_details_parts, author_parts)

    if patient_details_parts:
        sections['patient_details'] = '\n'.join(patient_details_parts)
    if author_parts:
        if 'signature' in sections:
            sections['signature'] += '\n\n' + '\n'.join(author_parts)
        else:
            sections['signature'] = '\n'.join(author_parts)

    # Parse paragraphs by heading keywords
    current_key = None
    current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if this paragraph is a heading
        is_heading = False
        # Bold or heading-style paragraphs (short text, often uppercase or styled)
        if para.style and para.style.name and 'heading' in para.style.name.lower():
            is_heading = True
        elif len(text) < 80 and text.endswith(':'):
            is_heading = True
        elif len(text) < 60 and text == text.upper() and len(text) > 5:
            is_heading = True
        elif para.runs and all(r.bold for r in para.runs if r.text.strip()):
            if len(text) < 100:
                is_heading = True

        if is_heading:
            # Try to match to GPR key
            lower = text.lower().rstrip(':').strip()
            matched_key = None
            for pattern, gpr_key in NARRATIVE_HEADING_MAP:
                if re.search(pattern, lower):
                    matched_key = gpr_key
                    break

            if matched_key:
                # Save previous section
                if current_key and current_content:
                    content = '\n\n'.join(current_content).strip()
                    if content:
                        if current_key in sections:
                            sections[current_key] += '\n\n' + content
                        else:
                            sections[current_key] = content
                current_key = matched_key
                current_content = []
                continue

        # Content paragraph
        if current_key:
            cleaned = _clean_content(text)
            if cleaned and not _is_question_text(cleaned):
                current_content.append(cleaned)

    # Save last section
    if current_key and current_content:
        content = '\n\n'.join(current_content).strip()
        if content:
            if current_key in sections:
                sections[current_key] += '\n\n' + content
            else:
                sections[current_key] = content

    return sections


# ============================================================
# Section Consolidation
# ============================================================

def _consolidate_to_gpr_keys(intermediate_sections: dict) -> dict:
    """Convert intermediate T131 keys to GPR card keys, merging with sub-headers."""
    gpr_sections = {}

    # Handle patient_details_raw and author_raw specially
    if 'patient_details_raw' in intermediate_sections:
        gpr_sections['patient_details'] = intermediate_sections.pop('patient_details_raw')
    if 'author_raw' in intermediate_sections:
        author = intermediate_sections.pop('author_raw')
        if 'signature' in gpr_sections:
            gpr_sections['signature'] += '\n\n' + author
        else:
            gpr_sections['signature'] = author

    # Handle sec_N fallback keys from numbered paragraphs
    for key in list(intermediate_sections.keys()):
        if key.startswith('sec_'):
            sec_num = int(key[4:])
            gpr_key = T131_TO_GPR.get(sec_num)
            if gpr_key:
                content = intermediate_sections.pop(key)
                label = T131_SUB_HEADERS.get(sec_num, f"Section {sec_num}")
                entry = f"**{label}:**\n{content}"
                if gpr_key in gpr_sections:
                    gpr_sections[gpr_key] += '\n\n' + entry
                else:
                    gpr_sections[gpr_key] = entry

    # Convert remaining intermediate keys
    for inter_key, content in intermediate_sections.items():
        if not content:
            continue
        # Skip keys already handled
        if inter_key in ('patient_details', 'patient_details_raw', 'author_raw'):
            continue

        gpr_key = _INTERMEDIATE_TO_GPR.get(inter_key)
        if not gpr_key:
            # Maybe it's already a GPR key (from narrative parser)
            valid_gpr_keys = {
                'patient_details', 'report_based_on', 'circumstances', 'background',
                'medical_history', 'psych_history', 'risk', 'substance_use',
                'forensic', 'medication', 'diagnosis', 'legal_criteria',
                'strengths', 'signature',
            }
            if inter_key in valid_gpr_keys:
                gpr_key = inter_key
            else:
                print(f"[GPR Parser] No GPR mapping for key '{inter_key}'")
                continue

        label = _INTERMEDIATE_LABELS.get(inter_key)

        if gpr_key in gpr_sections:
            if label:
                gpr_sections[gpr_key] += f'\n\n**{label}:**\n{content}'
            else:
                gpr_sections[gpr_key] += '\n\n' + content
        else:
            if label and inter_key != gpr_key:
                # Add sub-header only when merging multiple sections
                gpr_sections[gpr_key] = f'**{label}:**\n{content}'
            else:
                gpr_sections[gpr_key] = content

    # Post-process: if a GPR section only has one sub-section, remove the bold header
    for gpr_key, content in gpr_sections.items():
        bold_count = content.count('**')
        # If exactly one bold pair (opening + closing = 2 ** markers), it's a single sub-section
        if bold_count == 2 and content.startswith('**'):
            # Remove the single bold header since there's nothing to differentiate
            content = re.sub(r'^\*\*[^*]+\*\*:?\s*\n?', '', content).strip()
            gpr_sections[gpr_key] = content

    return gpr_sections


# ============================================================
# Main Entry Point
# ============================================================

def parse_gpr_report(file_path: str) -> dict:
    """Parse a medical tribunal report DOCX for GPR population.

    Returns:
        {
            "sections": {card_key: content_str, ...},
            "patient_info": {name, dob, gender, ...} or None,
            "format": "table_form" | "numbered_paragraphs" | "narrative",
            "source_file": filename
        }
        Returns empty dict on failure.
    """
    from docx import Document

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"[GPR Parser] Failed to open DOCX: {e}")
        return {}

    filename = os.path.basename(file_path)
    fmt = detect_format(doc)
    print(f"[GPR Parser] Detected format: {fmt} for {filename}")

    # Parse based on format
    if fmt == 'table_form':
        intermediate = _parse_table_form(doc)
    elif fmt == 'numbered_paragraphs':
        intermediate = _parse_numbered_paragraphs(doc)
    else:
        intermediate = _parse_narrative(doc)

    print(f"[GPR Parser] Intermediate keys: {list(intermediate.keys())}")

    if not intermediate:
        print(f"[GPR Parser] No sections found in {filename}")
        return {}

    # Consolidate to GPR card keys
    if fmt == 'narrative':
        # Narrative parser already returns GPR keys
        gpr_sections = intermediate
    else:
        gpr_sections = _consolidate_to_gpr_keys(intermediate)

    # Extract patient info
    patient_info = _extract_patient_info(gpr_sections)

    print(f"[GPR Parser] Parsed {len(gpr_sections)} GPR sections from {filename}")
    for key, content in gpr_sections.items():
        preview = content[:60].replace('\n', ' ')
        print(f"[GPR Parser]   {key}: {preview}...")

    return {
        "sections": gpr_sections,
        "patient_info": patient_info,
        "format": fmt,
        "source_file": filename,
    }
