# ================================================================
#  M2 FORM PAGE — Report Barring Discharge by Nearest Relative
#  Mental Health Act 1983 - Form M2 Regulation 25(1)(a) and (b)
#  Section 25 — Report barring discharge by nearest relative
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QPushButton, QFileDialog, QMessageBox, QToolButton,
    QRadioButton, QButtonGroup, QSpinBox, QComboBox, QCheckBox,
    QCompleter, QStyleFactory, QSlider
)

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    ICD10_DICT = load_icd10_dict()
except:
    ICD10_DICT = {}

from utils.resource_path import resource_path


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


class M2FormPage(QWidget):
    """Page for completing MHA Form M2 - Report Barring Discharge."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his", "self": "himself"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her", "self": "herself"}
        else:
            return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their", "self": "themselves"}

    def _generate_reasons_text(self):
        """Generate reasons text based on selected checkboxes (identical to A8)."""
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        # === PARAGRAPH 1: Demographics, Diagnosis, Nature/Degree ===
        para1_parts = []

        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")

        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            eth_simple = ethnicity.replace(" British", "").replace(" Other", "").replace("Mixed ", "")
            opening_parts.append(eth_simple)

        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        # Get diagnoses
        diagnoses = []
        if hasattr(self, 'dx_boxes') and self.dx_boxes:
            for combo in self.dx_boxes:
                meta = combo.currentData()
                if meta and isinstance(meta, dict):
                    dx = meta.get("diagnosis", "")
                    icd = meta.get("icd10", "")
                    if dx:
                        diagnoses.append(f"{dx} ({icd})" if icd else dx)

        if diagnoses:
            if opening_parts:
                demo_str = " ".join(opening_parts)
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {joined} which are mental disorders as defined by the Mental Health Act.")
            else:
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} suffers from {joined} which are mental disorders as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree to warrant detention for treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature to warrant detention for treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree to warrant detention for treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    nature_str = " and ".join(nature_types)
                    para1_parts.append(f"The nature of the illness is {nature_str}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        elif opening_parts:
            demo_str = " ".join(opening_parts)
            para1_parts.append(f"{name_display} is a {demo_str} who is currently detained under the Mental Health Act.")

        # === PARAGRAPH 2: Necessity (Health + Safety) ===
        para2_parts = []

        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("safety of others")

        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")

        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_reasons = []
            if self.poor_compliance_cb.isChecked():
                mh_reasons.append("non compliance")
            if self.limited_insight_cb.isChecked():
                mh_reasons.append("limited insight")

            if mh_reasons:
                reasons_str = "/".join(mh_reasons)
                para2_parts.append(f"Regarding health I would be concerned about {p['pos_l']} mental health deteriorating due to {reasons_str}.")
            else:
                para2_parts.append(f"Regarding health I would be concerned about {p['pos_l']} mental health deteriorating.")

        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health: {details}.")
            else:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health.")

        # Risk to self (A8 style - group by both/hist/curr)
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            risk_types = [
                ("self neglect", self.self_hist_neglect.isChecked(), self.self_curr_neglect.isChecked()),
                (f"placing of {p['self']} in risky situations", self.self_hist_risky.isChecked(), self.self_curr_risky.isChecked()),
                ("self harm", self.self_hist_harm.isChecked(), self.self_curr_harm.isChecked()),
            ]

            both_items = []
            hist_only = []
            curr_only = []

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            self_text = f"With respect to {p['pos_l']} own safety I am concerned about"
            parts = []
            if both_items:
                if len(both_items) == 1:
                    parts.append(f"historical and current {both_items[0]}")
                else:
                    parts.append(f"historical and current {', '.join(both_items[:-1])}, and of {both_items[-1]}")
            if hist_only:
                if len(hist_only) == 1:
                    parts.append(f"historical {hist_only[0]}")
                else:
                    parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
            if curr_only:
                if len(curr_only) == 1:
                    parts.append(f"current {curr_only[0]}")
                else:
                    parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")

            if parts:
                self_text += " " + ", and ".join(parts) + "."
            else:
                self_text += f" {p['pos_l']} safety."

            para2_parts.append(self_text)

        # Risk to others (A8 style)
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            risk_types = [
                ("violence to others", self.others_hist_violence.isChecked(), self.others_curr_violence.isChecked()),
                ("verbal aggression", self.others_hist_verbal.isChecked(), self.others_curr_verbal.isChecked()),
                ("sexual violence", self.others_hist_sexual.isChecked(), self.others_curr_sexual.isChecked()),
                ("stalking", self.others_hist_stalking.isChecked(), self.others_curr_stalking.isChecked()),
                ("arson", self.others_hist_arson.isChecked(), self.others_curr_arson.isChecked()),
            ]

            both_items = []
            hist_only = []
            curr_only = []

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            others_text = "With respect to risk to others I am concerned about the risk of"
            parts = []
            if both_items:
                parts.append(f"historical and current {', '.join(both_items)}")
            if hist_only:
                parts.append(f"historical {', '.join(hist_only)}")
            if curr_only:
                parts.append(f"current {', '.join(curr_only)}")

            if parts:
                others_text += " " + " and of ".join(parts) + "."
            else:
                others_text += f" {p['pos_l']} potential to cause harm."

            para2_parts.append(others_text)

        # === PARAGRAPH 3: Informal not indicated ===
        para3_parts = []

        if self.tried_failed_cb.isChecked():
            para3_parts.append("Previous attempts at informal admissions have not been successful and I would likewise be concerned about this recurring in this instance hence I do not believe informal admission currently would be appropriate.")

        if self.insight_cb.isChecked():
            para3_parts.append(f"{p['pos']} lack of insight is a significant concern and should {p['subj_l']} be discharged from section, I believe this would significantly impair {p['pos_l']} compliance if informal.")

        if self.compliance_cb.isChecked():
            if para3_parts:
                para3_parts.append(f"Compliance with treatment has also been a significant issue and I do not believe {p['subj_l']} would comply if informal.")
            else:
                para3_parts.append(f"Compliance with treatment has been a significant issue and I do not believe {p['subj_l']} would comply if informal.")

        if self.supervision_cb.isChecked():
            name = patient_name if patient_name else "the patient"
            para3_parts.append(f"I believe {name} needs careful community monitoring under the supervision afforded by the mental health act and I do not believe such supervision would be complied with should {p['subj_l']} remain in the community informally.")

        # Final statement for M2
        has_self_risk = self.safety_cb.isChecked() and self.self_harm_cb.isChecked()
        has_others_risk = self.safety_cb.isChecked() and self.others_cb.isChecked()
        if has_self_risk or has_others_risk:
            if has_self_risk and has_others_risk:
                danger = f"other persons and to {p['self']}"
            elif has_others_risk:
                danger = "other persons"
            else:
                danger = p['self']
            para3_parts.append(f"For these reasons, I am of the opinion that if {name_display} were to be discharged, {p['subj_l']} would be likely to act in a manner dangerous to {danger}.")

        paragraphs = []
        if para1_parts:
            paragraphs.append(" ".join(para1_parts))
        if para2_parts:
            paragraphs.append(" ".join(para2_parts))
        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        self.reasons.setPlainText("\n\n".join(paragraphs))

    # --- Control toggle handlers (identical to A8) ---
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._generate_reasons_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._generate_reasons_text()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._generate_reasons_text()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            self.self_hist_neglect.setChecked(False)
            self.self_hist_risky.setChecked(False)
            self.self_hist_harm.setChecked(False)
            self.self_curr_neglect.setChecked(False)
            self.self_curr_risky.setChecked(False)
            self.self_curr_harm.setChecked(False)
        self._generate_reasons_text()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            self.others_hist_violence.setChecked(False)
            self.others_hist_verbal.setChecked(False)
            self.others_hist_sexual.setChecked(False)
            self.others_hist_stalking.setChecked(False)
            self.others_hist_arson.setChecked(False)
            self.others_curr_violence.setChecked(False)
            self.others_curr_verbal.setChecked(False)
            self.others_curr_sexual.setChecked(False)
            self.others_curr_stalking.setChecked(False)
            self.others_curr_arson.setChecked(False)
        self._generate_reasons_text()

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #be185d; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form M2 — Report Barring Discharge by Nearest Relative")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        main_layout.addWidget(header)

        # Toolbar
        toolbar = QWidget()
        toolbar.setFixedHeight(60)
        toolbar.setStyleSheet("background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12);")
        tb_layout = QHBoxLayout(toolbar)
        tb_layout.setContentsMargins(16, 8, 16, 8)
        tb_layout.setSpacing(12)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(130, 38)
        export_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #be185d;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #9d174d; }
        """)
        export_btn.clicked.connect(self._export_docx)
        tb_layout.addWidget(export_btn)

        clear_btn = QToolButton()
        clear_btn.setText("Clear Form")
        clear_btn.setFixedSize(100, 38)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QToolButton {
                background: #ef4444;
                color: white;
                font-size: 13px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
            }
            QToolButton:hover { background: #dc2626; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        tb_layout.addWidget(clear_btn)
        tb_layout.addStretch()

        main_layout.addWidget(toolbar)

        # Scrollable form
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("background: #f9fafb;")

        form_container = QWidget()
        form_container.setStyleSheet("background: #f9fafb;")
        self.form_layout = QVBoxLayout(form_container)
        self.form_layout.setContentsMargins(40, 24, 40, 40)
        self.form_layout.setSpacing(24)

        self._build_part1()

        self.form_layout.addStretch()
        scroll.setWidget(form_container)
        main_layout.addWidget(scroll, 1)

    def _create_section_frame(self, title: str, color: str = "#be185d") -> QFrame:
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 12px; }")
        layout = QVBoxLayout(frame)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(16)
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet(f"font-size: 16px; font-weight: 700; color: {color};")
        layout.addWidget(title_lbl)
        return frame

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 13px;
            }
            QLineEdit:focus { border-color: #be185d; }
        """)
        return edit

    def _create_text_edit(self, placeholder: str = "", height: int = 100) -> QTextEdit:
        edit = QTextEdit()
        edit.setPlaceholderText(placeholder)
        edit.setMinimumHeight(height)
        edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #be185d; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return date_edit

    def _create_time_edit(self) -> QTimeEdit:
        time_edit = QTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("QTimeEdit { background: white; border: 1px solid #d1d5db; border-radius: 6px; padding: 10px 12px; font-size: 13px; }")
        return time_edit

    def _build_part1(self):
        frame = self._create_section_frame("PART 1 — Responsible Clinician")
        layout = frame.layout()

        # Hospital managers
        hosp_lbl = QLabel("To the managers of:")
        hosp_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #374151;")
        layout.addWidget(hosp_lbl)

        self.hospital = self._create_line_edit("Hospital name and address")
        layout.addWidget(self.hospital)

        # Nearest Relative Notice
        nr_lbl = QLabel("Notice of intention to discharge given by:")
        nr_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #374151; margin-top: 8px;")
        layout.addWidget(nr_lbl)

        self.nearest_relative = self._create_line_edit("Nearest relative name")
        layout.addWidget(self.nearest_relative)

        notice_row = QHBoxLayout()
        notice_row.setSpacing(20)

        time_lbl = QLabel("At time:")
        time_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        notice_row.addWidget(time_lbl)
        self.notice_time = self._create_time_edit()
        self.notice_time.setFixedWidth(100)
        notice_row.addWidget(self.notice_time)

        date_lbl = QLabel("On date:")
        date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        notice_row.addWidget(date_lbl)
        self.notice_date = self._create_date_edit()
        self.notice_date.setFixedWidth(140)
        notice_row.addWidget(self.notice_date)

        notice_row.addStretch()
        layout.addLayout(notice_row)

        # Patient
        pt_lbl = QLabel("Intention to discharge:")
        pt_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #374151; margin-top: 8px;")
        layout.addWidget(pt_lbl)

        self.patient_name = self._create_line_edit("Patient full name")
        layout.addWidget(self.patient_name)

        # Patient Demographics row
        demo_row = QHBoxLayout()
        demo_row.setSpacing(16)

        # Age
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        demo_row.addWidget(age_lbl)

        self.age_spin = QSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setFixedWidth(70)
        self.age_spin.setStyleSheet("QSpinBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 13px; }")
        demo_row.addWidget(self.age_spin)

        demo_row.addSpacing(20)

        # Gender
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("Male")
        self.gender_female = QRadioButton("Female")
        self.gender_other = QRadioButton("Other")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("font-size: 13px;")
            self.gender_group.addButton(rb)
            demo_row.addWidget(rb)

        demo_row.addSpacing(20)

        # Ethnicity (same as A8)
        self.ethnicity_combo = QComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems([
            "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race", "Not specified"
        ])
        self.ethnicity_combo.setFixedWidth(160)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 13px; }")
        demo_row.addWidget(self.ethnicity_combo)

        demo_row.addStretch()
        layout.addLayout(demo_row)

        # Opinion
        op_lbl = QLabel("I am of the opinion that the patient, if discharged, would be likely to act in a manner dangerous to other persons or to himself/herself.")
        op_lbl.setWordWrap(True)
        op_lbl.setStyleSheet("font-size: 13px; color: #374151; padding: 8px; background: #fef2f2; border-radius: 6px; margin-top: 8px;")
        layout.addWidget(op_lbl)

        # Reasons - split layout with controls on right
        reasons_lbl = QLabel("Reasons for my opinion:")
        reasons_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #374151; margin-top: 8px;")
        layout.addWidget(reasons_lbl)

        split_layout = QHBoxLayout()
        split_layout.setSpacing(20)

        # Left: Text area
        left_container = QWidget()
        left_layout = QVBoxLayout(left_container)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.setSpacing(8)

        info = QLabel("Click options on the right to auto-generate text:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 11px; color: #6b7280; padding: 6px; background: #fce7f3; border-radius: 4px;")
        left_layout.addWidget(info)

        self.reasons = QTextEdit()
        self.reasons.setPlaceholderText("Reasons will be generated here...")
        self.reasons.setMinimumHeight(200)
        self.reasons.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 12px;
                font-size: 13px;
            }
            QTextEdit:focus { border-color: #be185d; }
        """)
        left_layout.addWidget(self.reasons)
        split_layout.addWidget(left_container, 3)

        # Right: Controls panel (identical to A8)
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        right_scroll.setFixedWidth(400)

        right_container = QWidget()
        right_container.setFixedWidth(380)
        right_container.setStyleSheet("background: transparent;")
        right_layout = QVBoxLayout(right_container)
        right_layout.setContentsMargins(0, 0, 8, 0)
        right_layout.setSpacing(12)

        # --- Mental Disorder (ICD-10) - 2 boxes ---
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: 1px solid #bbf7d0; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(8)

        md_header = QLabel("Mental Disorder")
        md_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        self.dx_boxes = []
        for i in range(2):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText("Primary diagnosis..." if i == 0 else "Secondary (optional)...")

            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(ICD10_DICT.items(), key=lambda x: x[0].lower()):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})

            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(12)
            combo.setStyleSheet("""
                QComboBox { padding: 6px; font-size: 12px; border: 1px solid #d1d5db; border-radius: 4px; background: white; }
                QComboBox QAbstractItemView { min-width: 300px; }
            """)
            combo.currentIndexChanged.connect(self._generate_reasons_text)
            md_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        right_layout.addWidget(md_frame)

        # --- Legal Criteria ---
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(6)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature with sub-options
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._generate_reasons_text)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._generate_reasons_text)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._generate_reasons_text)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree with slider
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(120)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 12px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 12px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._generate_reasons_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity section
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 12px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health with sub-options
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._generate_reasons_text)
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._generate_reasons_text)
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 12px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 12px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._generate_reasons_text)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety with sub-options
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # === SELF SECTION ===
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Self placement in risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_hist_harm)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Self placement in risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # === OTHERS SECTION ===
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 12px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence to others")
        self.others_hist_violence.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_verbal)

        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_sexual.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_sexual.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_sexual)

        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_stalking.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_stalking.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_stalking)

        self.others_hist_arson = QCheckBox("Arson")
        self.others_hist_arson.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_hist_arson.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_arson)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 12px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence to others")
        self.others_curr_violence.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_sexual.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_sexual.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_sexual)

        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_stalking.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_stalking.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_stalking)

        self.others_curr_arson = QCheckBox("Arson")
        self.others_curr_arson.setStyleSheet("font-size: 12px; color: #9ca3af;")
        self.others_curr_arson.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_arson)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        right_layout.addWidget(lc_frame)

        # --- Informal Not Indicated ---
        inf_frame = QFrame()
        inf_frame.setStyleSheet("QFrame { background: #fef2f2; border: 1px solid #fecaca; border-radius: 8px; }")
        inf_layout = QVBoxLayout(inf_frame)
        inf_layout.setContentsMargins(12, 10, 12, 10)
        inf_layout.setSpacing(4)

        inf_header = QLabel("Informal Not Indicated")
        inf_header.setStyleSheet("font-size: 12px; font-weight: 700; color: #991b1b;")
        inf_layout.addWidget(inf_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed")
        self.tried_failed_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.tried_failed_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.tried_failed_cb)

        self.insight_cb = QCheckBox("Lack of Insight")
        self.insight_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.insight_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.insight_cb)

        self.compliance_cb = QCheckBox("Compliance Issues")
        self.compliance_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.compliance_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.compliance_cb)

        self.supervision_cb = QCheckBox("Needs MHA Supervision")
        self.supervision_cb.setStyleSheet("font-size: 12px; color: #374151;")
        self.supervision_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.supervision_cb)

        right_layout.addWidget(inf_frame)
        right_layout.addStretch()

        right_scroll.setWidget(right_container)
        split_layout.addWidget(right_scroll)

        layout.addLayout(split_layout)

        # RC Details
        rc_lbl = QLabel("Responsible Clinician")
        rc_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #374151; margin-top: 8px;")
        layout.addWidget(rc_lbl)

        row1 = QHBoxLayout()
        row1.setSpacing(12)
        self.rc_name = self._create_line_edit("RC full name")
        row1.addWidget(self.rc_name, 1)
        self.rc_email = self._create_line_edit("Email (if applicable)")
        row1.addWidget(self.rc_email, 1)
        layout.addLayout(row1)

        # Furnishing Report Section
        furnish_lbl = QLabel("I am furnishing this report by:")
        furnish_lbl.setStyleSheet("font-size: 13px; font-weight: 600; color: #374151; margin-top: 12px;")
        layout.addWidget(furnish_lbl)

        self.furnish_group = QButtonGroup(self)

        # Option 1: Internal mail with time
        opt1_row = QHBoxLayout()
        opt1_row.setSpacing(8)
        self.furnish_internal_mail = QRadioButton("consigning it to the hospital managers' internal mail system today at")
        self.furnish_internal_mail.setChecked(True)
        self.furnish_internal_mail.setStyleSheet("font-size: 12px;")
        self.furnish_group.addButton(self.furnish_internal_mail)
        opt1_row.addWidget(self.furnish_internal_mail)
        self.furnish_time = self._create_time_edit()
        self.furnish_time.setFixedWidth(80)
        opt1_row.addWidget(self.furnish_time)
        opt1_row.addStretch()
        layout.addLayout(opt1_row)

        # Option 2: Electronic communication
        self.furnish_electronic = QRadioButton("today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication")
        self.furnish_electronic.setStyleSheet("font-size: 12px;")
        self.furnish_group.addButton(self.furnish_electronic)
        layout.addWidget(self.furnish_electronic)

        # Option 3: Other delivery
        self.furnish_other = QRadioButton("sending or delivering it without using the hospital managers' internal mail system")
        self.furnish_other.setStyleSheet("font-size: 12px;")
        self.furnish_group.addButton(self.furnish_other)
        layout.addWidget(self.furnish_other)

        # Signature
        sig_row = QHBoxLayout()
        sig_row.setSpacing(20)

        sig_date_lbl = QLabel("Date:")
        sig_date_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_date_lbl)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(140)
        sig_row.addWidget(self.rc_sig_date)

        sig_time_lbl = QLabel("Time:")
        sig_time_lbl.setStyleSheet("font-size: 13px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_time_lbl)
        self.rc_sig_time = self._create_time_edit()
        self.rc_sig_time.setFixedWidth(100)
        sig_row.addWidget(self.rc_sig_time)

        sig_row.addStretch()
        layout.addLayout(sig_row)

        self.form_layout.addWidget(frame)

    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.hospital.clear()
            self.nearest_relative.clear()
            self.notice_time.setTime(QTime.currentTime())
            self.notice_date.setDate(QDate.currentDate())
            self.patient_name.clear()
            self.age_spin.setValue(0)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.ethnicity_combo.setCurrentIndex(0)
            # Mental Disorder (ICD-10)
            for combo in self.dx_boxes:
                combo.setCurrentIndex(0)
            # Legal criteria (A8-style)
            self.nature_cb.setChecked(False)
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.degree_slider.setValue(2)
            self.degree_details.clear()
            self.health_cb.setChecked(False)
            self.mental_health_cb.setChecked(False)
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
            self.physical_health_details.clear()
            self.safety_cb.setChecked(False)
            self.self_harm_cb.setChecked(False)
            self.self_hist_neglect.setChecked(False)
            self.self_hist_risky.setChecked(False)
            self.self_hist_harm.setChecked(False)
            self.self_curr_neglect.setChecked(False)
            self.self_curr_risky.setChecked(False)
            self.self_curr_harm.setChecked(False)
            self.others_cb.setChecked(False)
            self.others_hist_violence.setChecked(False)
            self.others_hist_verbal.setChecked(False)
            self.others_hist_sexual.setChecked(False)
            self.others_hist_stalking.setChecked(False)
            self.others_hist_arson.setChecked(False)
            self.others_curr_violence.setChecked(False)
            self.others_curr_verbal.setChecked(False)
            self.others_curr_sexual.setChecked(False)
            self.others_curr_stalking.setChecked(False)
            self.others_curr_arson.setChecked(False)
            # Informal not indicated
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.reasons.clear()
            self.rc_name.clear()
            self.rc_email.clear()
            self.furnish_internal_mail.setChecked(True)
            self.furnish_time.setTime(QTime.currentTime())
            self.rc_sig_date.setDate(QDate.currentDate())
            self.rc_sig_time.setTime(QTime.currentTime())

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form M2",
            f"Form_M2_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_M2_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form M2 template not found.")
                return

            doc = Document(template_path)

            def set_para_text(para, new_text):
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

            def highlight_yellow(para):
                for run in para.runs:
                    rPr = run._element.get_or_add_rPr()
                    shd = rPr.find(qn('w:shd'))
                    if shd is None:
                        shd = OxmlElement('w:shd')
                        rPr.append(shd)
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFFCC')

            def add_strikethrough(run):
                """Add strikethrough to a run."""
                rPr = run._element.get_or_add_rPr()
                strike = OxmlElement('w:strike')
                strike.set(qn('w:val'), 'true')
                rPr.append(strike)

            def highlight_run_yellow(run):
                """Add yellow highlight to a single run."""
                rPr = run._element.get_or_add_rPr()
                shd = rPr.find(qn('w:shd'))
                if shd is None:
                    shd = OxmlElement('w:shd')
                    rPr.append(shd)
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFFCC')

            paragraphs = doc.paragraphs

            # Hospital (para 5)
            if self.hospital.text().strip():
                set_para_text(paragraphs[5], self.hospital.text())
                highlight_yellow(paragraphs[5])

            # Nearest relative (para 7)
            if self.nearest_relative.text().strip():
                set_para_text(paragraphs[7], self.nearest_relative.text())
                highlight_yellow(paragraphs[7])

            # Notice time (para 9)
            notice_time = self.notice_time.time().toString("HH:mm")
            set_para_text(paragraphs[9], notice_time)
            highlight_yellow(paragraphs[9])

            # Notice date (para 11)
            notice_date = self.notice_date.date().toString("dd MMMM yyyy")
            set_para_text(paragraphs[11], notice_date)
            highlight_yellow(paragraphs[11])

            # Patient name (para 13)
            if self.patient_name.text().strip():
                set_para_text(paragraphs[13], self.patient_name.text())
                highlight_yellow(paragraphs[13])

            # Reasons (para 16-18)
            if self.reasons.toPlainText().strip():
                set_para_text(paragraphs[16], self.reasons.toPlainText())
                highlight_yellow(paragraphs[16])

            # RC name (para 26) - keep "PRINT NAME" label, add name in yellow (no brackets)
            if "PRINT NAME" in paragraphs[26].text:
                for run in paragraphs[26].runs:
                    run.text = ""
                paragraphs[26].add_run("PRINT NAME  ")
                name_val = self.rc_name.text().strip() if self.rc_name.text().strip() else "                    "
                name_run = paragraphs[26].add_run(name_val)
                highlight_run_yellow(name_run)

            # RC email (para 27) - keep label, add email in yellow (no brackets)
            if "Email address" in paragraphs[27].text:
                for run in paragraphs[27].runs:
                    run.text = ""
                paragraphs[27].add_run("Email address (if applicable)  ")
                email_val = self.rc_email.text().strip() if self.rc_email.text().strip() else "                    "
                email_run = paragraphs[27].add_run(email_val)
                highlight_run_yellow(email_run)

            # RC sig date and time (para 28) - keep labels, add values in yellow (no brackets)
            sig_date = self.rc_sig_date.date().toString("dd MMMM yyyy")
            sig_time = self.rc_sig_time.time().toString("HH:mm")
            if "Date" in paragraphs[28].text:
                for run in paragraphs[28].runs:
                    run.text = ""
                paragraphs[28].add_run("Date  ")
                date_run = paragraphs[28].add_run(sig_date)
                highlight_run_yellow(date_run)
                paragraphs[28].add_run("                                                 Time  ")
                time_run = paragraphs[28].add_run(sig_time)
                highlight_run_yellow(time_run)

            # Furnishing report options (paras 21, 23, 24)
            furnish_time = self.furnish_time.time().toString("HH:mm")

            # Para 21: Internal mail with time
            if "internal mail system" in paragraphs[21].text:
                for run in paragraphs[21].runs:
                    run.text = ""
                if self.furnish_internal_mail.isChecked():
                    run1 = paragraphs[21].add_run("consigning it to the hospital managers' internal mail system today at ")
                    run2 = paragraphs[21].add_run(furnish_time)
                    highlight_run_yellow(run2)
                    run3 = paragraphs[21].add_run(". ")
                else:
                    run1 = paragraphs[21].add_run("consigning it to the hospital managers' internal mail system today at [time]. ")
                    add_strikethrough(run1)

            # Para 23: Electronic communication
            if "electronic communication" in paragraphs[23].text or "today sending" in paragraphs[23].text:
                for run in paragraphs[23].runs:
                    run.text = ""
                run = paragraphs[23].add_run("today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.")
                if not self.furnish_electronic.isChecked():
                    add_strikethrough(run)

            # Para 24: Other delivery
            if "without using" in paragraphs[24].text:
                for run in paragraphs[24].runs:
                    run.text = ""
                run = paragraphs[24].add_run("sending or delivering it without using the hospital managers' internal mail system.")
                if not self.furnish_other.isChecked():
                    add_strikethrough(run)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form M2 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    def get_state(self) -> dict:
        # Determine furnish method
        if self.furnish_internal_mail.isChecked():
            furnish_method = "internal_mail"
        elif self.furnish_electronic.isChecked():
            furnish_method = "electronic"
        else:
            furnish_method = "other"

        # Determine gender
        gender = None
        if self.gender_male.isChecked():
            gender = "male"
        elif self.gender_female.isChecked():
            gender = "female"
        elif self.gender_other.isChecked():
            gender = "other"

        return {
            "hospital": self.hospital.text(),
            "nearest_relative": self.nearest_relative.text(),
            "notice_time": self.notice_time.time().toString("HH:mm"),
            "notice_date": self.notice_date.date().toString("yyyy-MM-dd"),
            "patient_name": self.patient_name.text(),
            "age": self.age_spin.value(),
            "gender": gender,
            "ethnicity": self.ethnicity_combo.currentText(),
            # Mental Disorder (ICD-10)
            "diagnosis_primary": self.dx_boxes[0].currentText() if self.dx_boxes else "",
            "diagnosis_secondary": self.dx_boxes[1].currentText() if len(self.dx_boxes) > 1 else "",
            # Legal criteria (A8-style)
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_slider": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_hist_neglect": self.self_hist_neglect.isChecked(),
            "self_hist_risky": self.self_hist_risky.isChecked(),
            "self_hist_harm": self.self_hist_harm.isChecked(),
            "self_curr_neglect": self.self_curr_neglect.isChecked(),
            "self_curr_risky": self.self_curr_risky.isChecked(),
            "self_curr_harm": self.self_curr_harm.isChecked(),
            "others": self.others_cb.isChecked(),
            "others_hist_violence": self.others_hist_violence.isChecked(),
            "others_hist_verbal": self.others_hist_verbal.isChecked(),
            "others_hist_sexual": self.others_hist_sexual.isChecked(),
            "others_hist_stalking": self.others_hist_stalking.isChecked(),
            "others_hist_arson": self.others_hist_arson.isChecked(),
            "others_curr_violence": self.others_curr_violence.isChecked(),
            "others_curr_verbal": self.others_curr_verbal.isChecked(),
            "others_curr_sexual": self.others_curr_sexual.isChecked(),
            "others_curr_stalking": self.others_curr_stalking.isChecked(),
            "others_curr_arson": self.others_curr_arson.isChecked(),
            # Informal not indicated
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "reasons": self.reasons.toPlainText(),
            "rc_name": self.rc_name.text(),
            "rc_email": self.rc_email.text(),
            "furnish_method": furnish_method,
            "furnish_time": self.furnish_time.time().toString("HH:mm"),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "rc_sig_time": self.rc_sig_time.time().toString("HH:mm"),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.hospital.setText(state.get("hospital", ""))
        self.nearest_relative.setText(state.get("nearest_relative", ""))
        if state.get("notice_time"):
            self.notice_time.setTime(QTime.fromString(state["notice_time"], "HH:mm"))
        if state.get("notice_date"):
            self.notice_date.setDate(QDate.fromString(state["notice_date"], "yyyy-MM-dd"))
        self.patient_name.setText(state.get("patient_name", ""))
        self.age_spin.setValue(state.get("age", 0))
        # Restore gender
        gender = state.get("gender")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        # Restore ethnicity
        ethnicity = state.get("ethnicity", "Ethnicity")
        idx = self.ethnicity_combo.findText(ethnicity)
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        # Restore diagnosis (ICD-10)
        if self.dx_boxes:
            dx_primary = state.get("diagnosis_primary", "Not specified")
            idx = self.dx_boxes[0].findText(dx_primary)
            if idx >= 0:
                self.dx_boxes[0].setCurrentIndex(idx)
            if len(self.dx_boxes) > 1:
                dx_secondary = state.get("diagnosis_secondary", "Not specified")
                idx = self.dx_boxes[1].findText(dx_secondary)
                if idx >= 0:
                    self.dx_boxes[1].setCurrentIndex(idx)
        # Restore legal criteria (A8-style)
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_slider", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_hist_neglect.setChecked(state.get("self_hist_neglect", False))
        self.self_hist_risky.setChecked(state.get("self_hist_risky", False))
        self.self_hist_harm.setChecked(state.get("self_hist_harm", False))
        self.self_curr_neglect.setChecked(state.get("self_curr_neglect", False))
        self.self_curr_risky.setChecked(state.get("self_curr_risky", False))
        self.self_curr_harm.setChecked(state.get("self_curr_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.others_hist_violence.setChecked(state.get("others_hist_violence", False))
        self.others_hist_verbal.setChecked(state.get("others_hist_verbal", False))
        self.others_hist_sexual.setChecked(state.get("others_hist_sexual", False))
        self.others_hist_stalking.setChecked(state.get("others_hist_stalking", False))
        self.others_hist_arson.setChecked(state.get("others_hist_arson", False))
        self.others_curr_violence.setChecked(state.get("others_curr_violence", False))
        self.others_curr_verbal.setChecked(state.get("others_curr_verbal", False))
        self.others_curr_sexual.setChecked(state.get("others_curr_sexual", False))
        self.others_curr_stalking.setChecked(state.get("others_curr_stalking", False))
        self.others_curr_arson.setChecked(state.get("others_curr_arson", False))
        # Restore informal not indicated
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        self.reasons.setPlainText(state.get("reasons", ""))
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_email.setText(state.get("rc_email", ""))
        # Restore furnishing method
        furnish_method = state.get("furnish_method", "internal_mail")
        if furnish_method == "electronic":
            self.furnish_electronic.setChecked(True)
        elif furnish_method == "other":
            self.furnish_other.setChecked(True)
        else:
            self.furnish_internal_mail.setChecked(True)
        if state.get("furnish_time"):
            self.furnish_time.setTime(QTime.fromString(state["furnish_time"], "HH:mm"))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        if state.get("rc_sig_time"):
            self.rc_sig_time.setTime(QTime.fromString(state["rc_sig_time"], "HH:mm"))
