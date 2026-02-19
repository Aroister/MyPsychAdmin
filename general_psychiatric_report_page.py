# ================================================================
#  GENERAL PSYCHIATRIC REPORT PAGE
#  Based on CPA REPORT MEDICAL template format
# ================================================================

from __future__ import annotations

import os
import re
import sys
from datetime import datetime
from typing import Optional

from PySide6.QtCore import Qt, Signal, QSize, QEvent
from PySide6.QtGui import QColor, QFontDatabase
from PySide6.QtWidgets import QGraphicsDropShadowEffect
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QSplitter, QStackedWidget, QTextEdit,
    QSizePolicy, QPushButton, QToolButton, QComboBox, QColorDialog,
    QLineEdit, QDateEdit, QFormLayout, QTableWidget, QTableWidgetItem,
    QHeaderView, QCheckBox, QSpinBox, QFileDialog, QMessageBox,
    QGridLayout, QApplication
)
from PySide6.QtCore import QDate
from mypsy_richtext_editor import MyPsychAdminRichTextEditor

from background_history_popup import BackgroundHistoryPopup, CollapsibleSection, ResizableSection
from shared_widgets import create_zoom_row, add_lock_to_popup
from physical_health_popup import PhysicalHealthPopup
from drugs_alcohol_popup import DrugsAlcoholPopup


# ================================================================
# NO-WHEEL WIDGETS (prevents scroll from changing value)
# ================================================================
from PySide6.QtWidgets import QSlider
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelComboBox(QComboBox):
    """ComboBox that ignores wheel events to prevent accidental changes while scrolling."""
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# GPR TOOLBAR
# ================================================================

class GPRToolbar(QWidget):
    """Toolbar for the General Psychiatric Report Page."""

    # Formatting signals
    set_font_family = Signal(str)
    set_font_size = Signal(int)

    toggle_bold = Signal()
    toggle_italic = Signal()
    toggle_underline = Signal()

    set_text_color = Signal(QColor)
    set_highlight_color = Signal(QColor)

    undo = Signal()
    redo = Signal()

    export_docx = Signal()
    check_spelling = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(80)
        self.setStyleSheet("""
            GPRToolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 18px;
                font-weight: 500;
            }
            QToolButton:hover {
                background: rgba(0,0,0,0.08);
            }
            QComboBox {
                background: rgba(255,255,255,0.85);
                color: #333333;
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 17px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #e0e0e0;
            }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(False)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        container.setMinimumWidth(1200)  # Force scrollbar when viewport is smaller
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 8, 2)
        layout.setSpacing(10)

        # Export DOCX button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(160, 42)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 18px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #1d4ed8; }
            QToolButton:pressed { background: #1e40af; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Uploaded Docs button (dropdown menu)
        from PySide6.QtWidgets import QMenu
        import_btn = QToolButton()
        import_btn.setText("Uploaded Docs")
        import_btn.setFixedSize(160, 42)
        import_btn.setPopupMode(QToolButton.InstantPopup)
        self.upload_menu = QMenu()
        import_btn.setMenu(self.upload_menu)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 18px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover { background: #059669; }
            QToolButton:pressed { background: #047857; }
            QToolButton::menu-indicator { image: none; }
        """)
        layout.addWidget(import_btn)

        # Font Family
        self.font_combo = NoWheelComboBox()
        self.font_combo.setFixedWidth(160)
        families = QFontDatabase.families()
        if sys.platform == "win32":
            preferred = ["Segoe UI", "Calibri", "Cambria", "Arial", "Times New Roman"]
        else:
            preferred = ["Avenir Next", "Avenir", "SF Pro Text", "Helvetica Neue", "Helvetica"]
        added = set()
        for f in preferred:
            if f in families:
                self.font_combo.addItem(f)
                added.add(f)
        for f in families:
            if f not in added:
                self.font_combo.addItem(f)
        self.font_combo.currentTextChanged.connect(self.set_font_family.emit)
        layout.addWidget(self.font_combo)

        # Font Size
        self.size_combo = NoWheelComboBox()
        self.size_combo.setFixedWidth(60)
        for sz in [8, 9, 10, 11, 12, 14, 16, 18, 20]:
            self.size_combo.addItem(str(sz))
        self.size_combo.currentTextChanged.connect(lambda v: self.set_font_size.emit(int(v)))
        layout.addWidget(self.size_combo)

        def btn(label, slot):
            b = QToolButton()
            b.setText(label)
            b.setMinimumWidth(36)
            b.clicked.connect(slot)
            return b

        layout.addWidget(btn("B", self.toggle_bold.emit))
        layout.addWidget(btn("I", self.toggle_italic.emit))
        layout.addWidget(btn("U", self.toggle_underline.emit))
        layout.addWidget(btn("A", self._choose_text_color))
        layout.addWidget(btn("⟲", self.undo.emit))
        layout.addWidget(btn("⟳", self.redo.emit))

        # Spell Check button
        spell_btn = QToolButton()
        spell_btn.setText("Spell Check")
        spell_btn.setFixedSize(120, 38)
        spell_btn.setStyleSheet("""
            QToolButton {
                background: #f59e0b;
                color: white;
                font-size: 16px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 6px 12px;
            }
            QToolButton:hover { background: #d97706; }
            QToolButton:pressed { background: #b45309; }
        """)
        spell_btn.setToolTip("Jump to next spelling error")
        spell_btn.clicked.connect(self.check_spelling.emit)
        layout.addWidget(spell_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)

    def _choose_text_color(self):
        col = QColorDialog.getColor(QColor("black"), self)
        if col.isValid():
            self.set_text_color.emit(col)


# ================================================================
# CARD WIDGET
# ================================================================

class GPRCardWidget(QFrame):
    """A clickable card for a report section."""

    clicked = Signal(str)

    STYLE_NORMAL = """
        GPRCardWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
            padding: 16px;
        }
        GPRCardWidget:hover {
            border-color: #2563eb;
            background: #eff6ff;
        }
    """

    STYLE_SELECTED = """
        GPRCardWidget {
            background: #dbeafe;
            border: 2px solid #2563eb;
            border-left: 4px solid #1d4ed8;
            border-radius: 12px;
            padding: 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        header_row.addWidget(title_lbl)
        header_row.addStretch()

        layout.addLayout(header_row)

        # Editor (rich text with formatting support)
        self.editor = MyPsychAdminRichTextEditor()
        self.editor.setPlaceholderText("Click to edit...")
        self.editor.setReadOnly(False)
        self._editor_height = 100
        self.editor.setMinimumHeight(60)
        self.editor.setMaximumHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
                color: #374151;
            }
        """)
        layout.addWidget(self.editor)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.editor, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(12)
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #2563eb;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(60, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setMinimumHeight(self._editor_height)
                self.editor.setMaximumHeight(self._editor_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        return self._selected


# ================================================================
# PATIENT DETAILS POPUP
# ================================================================

class GPRPatientDetailsPopup(QWidget):
    """Popup for entering patient details."""

    sent = Signal(str)
    gender_changed = Signal(str)  # Emitted when gender selection changes
    age_changed = Signal(int)  # Emitted when patient age changes

    def __init__(self, parent=None):
        super().__init__(parent)
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Create scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        # Content widget inside scroll area
        content = QWidget()
        layout = QVBoxLayout(content)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        title = QLabel("Patient Details")
        title.setStyleSheet("font-size: 20px; font-weight: 700; color: #2563eb;")
        layout.addWidget(title)

        # Date picker styling to avoid black background
        date_picker_style = """
            QDateEdit {
                background: white;
                color: #333333;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 10px 8px;
                min-height: 24px;
                font-size: 18px;
            }
            QDateEdit::drop-down {
                border: none;
                width: 20px;
            }
            QCalendarWidget {
                background: white;
            }
            QCalendarWidget QWidget {
                background: white;
                color: #333333;
            }
            QCalendarWidget QToolButton {
                background: #f3f4f6;
                color: #333333;
            }
            QCalendarWidget QMenu {
                background: white;
                color: #333333;
            }
            QCalendarWidget QSpinBox {
                background: white;
                color: #333333;
            }
        """

        # Form layout for patient details
        form = QFormLayout()
        form.setSpacing(20)
        form.setVerticalSpacing(20)

        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("Patient name")
        self.name_edit.setStyleSheet("QLineEdit { font-size: 19px; padding: 10px 8px; min-height: 24px; }")
        form.addRow("Name:", self.name_edit)

        # Gender dropdown
        self.gender_combo = NoWheelComboBox()
        self.gender_combo.addItems(["", "Male", "Female", "Non-binary", "Other"])
        self.gender_combo.setStyleSheet("""
            QComboBox {
                background: white;
                color: #333333;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 10px 8px;
                padding-right: 25px;
                min-width: 120px;
                min-height: 24px;
                font-size: 18px;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 24px;
                border-left: 1px solid #d1d5db;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
                background: #f3f4f6;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 6px solid #374151;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #dbeafe;
                border: 1px solid #d1d5db;
                font-size: 18px;
            }
        """)
        form.addRow("Gender:", self.gender_combo)

        self.dob_edit = QDateEdit()
        self.dob_edit.setCalendarPopup(True)
        self.dob_edit.setDate(QDate(1980, 1, 1))
        self.dob_edit.setStyleSheet(date_picker_style)
        form.addRow("Date of Birth:", self.dob_edit)

        # Age - read-only, auto-calculated from DOB
        self.age_label = QLineEdit()
        self.age_label.setReadOnly(True)
        self.age_label.setStyleSheet("""
            QLineEdit {
                background: #f3f4f6;
                color: #374151;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 10px 8px;
                min-height: 24px;
                font-size: 18px;
            }
        """)
        form.addRow("Age:", self.age_label)

        # Section dropdown with MHA options
        self.section_combo = NoWheelComboBox()
        self.section_combo.setEditable(True)  # Allow custom entry
        self.section_combo.addItems([
            "",
            "Section 2",
            "Section 3",
            "Section 37",
            "Section 37/41",
            "Section 47",
            "Section 47/49",
            "Section 48/49",
            "CTO",
            "Conditional Discharge",
            "Informal",
            "Other"
        ])
        self.section_combo.setStyleSheet("""
            QComboBox {
                background: white;
                color: #333333;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 10px 8px;
                padding-right: 25px;
                min-width: 150px;
                min-height: 24px;
                font-size: 18px;
            }
            QComboBox::drop-down {
                subcontrol-origin: padding;
                subcontrol-position: center right;
                width: 24px;
                border-left: 1px solid #d1d5db;
                border-top-right-radius: 4px;
                border-bottom-right-radius: 4px;
                background: #f3f4f6;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
                image: none;
                border-left: 4px solid transparent;
                border-right: 4px solid transparent;
                border-top: 6px solid #374151;
            }
            QComboBox::drop-down:hover {
                background: #e5e7eb;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #dbeafe;
                border: 1px solid #d1d5db;
                font-size: 18px;
            }
        """)
        form.addRow("Section (MHA):", self.section_combo)

        self.admission_date_edit = QDateEdit()
        self.admission_date_edit.setCalendarPopup(True)
        self.admission_date_edit.setStyleSheet(date_picker_style)
        self.admission_date_edit.setDate(QDate.currentDate())
        form.addRow("Admission Date:", self.admission_date_edit)

        self.location_edit = QLineEdit()
        self.location_edit.setPlaceholderText("Hospital/Ward")
        self.location_edit.setStyleSheet("QLineEdit { font-size: 19px; padding: 10px 8px; min-height: 24px; }")
        form.addRow("Current Location:", self.location_edit)

        self.report_by_edit = QLineEdit()
        self.report_by_edit.setPlaceholderText("Clinician name")
        self.report_by_edit.setStyleSheet("QLineEdit { font-size: 19px; padding: 10px 8px; min-height: 24px; }")
        form.addRow("Report By:", self.report_by_edit)

        self.date_seen_edit = QDateEdit()
        self.date_seen_edit.setCalendarPopup(True)
        self.date_seen_edit.setDate(QDate.currentDate())
        self.date_seen_edit.setStyleSheet(date_picker_style)
        form.addRow("Date Seen:", self.date_seen_edit)

        layout.addLayout(form)

        # Auto-calculate age when DOB changes
        self.dob_edit.dateChanged.connect(self._update_age)
        # Initialize age
        self._update_age(self.dob_edit.date())

        # Auto-sync fields to card when they change
        self.name_edit.textChanged.connect(self._auto_sync_to_card)
        self.gender_combo.currentTextChanged.connect(self._auto_sync_to_card)
        self.gender_combo.currentTextChanged.connect(self.gender_changed.emit)  # Notify parent of gender changes
        self.dob_edit.dateChanged.connect(self._auto_sync_to_card)
        self.section_combo.currentTextChanged.connect(self._auto_sync_to_card)
        self.admission_date_edit.dateChanged.connect(self._auto_sync_to_card)
        self.location_edit.textChanged.connect(self._auto_sync_to_card)
        self.report_by_edit.textChanged.connect(self._auto_sync_to_card)
        self.date_seen_edit.dateChanged.connect(self._auto_sync_to_card)

        # Style form labels
        for i in range(form.rowCount()):
            label_item = form.itemAt(i, QFormLayout.ItemRole.LabelRole)
            if label_item and label_item.widget():
                label_item.widget().setStyleSheet("font-size: 18px; color: #374151;")

        layout.addStretch()

        # Add content to scroll area
        scroll.setWidget(content)
        main_layout.addWidget(scroll)

    def _update_age(self, date):
        today = QDate.currentDate()
        age = today.year() - date.year()
        if today.month() < date.month() or (today.month() == date.month() and today.day() < date.day()):
            age -= 1
        age = max(0, age)
        self.age_label.setText(str(age))
        self.age_changed.emit(age)

    def _auto_sync_to_card(self):
        """Auto-sync all fields to card when any field changes."""
        self._send_to_report()

    def _send_to_report(self):
        lines = []
        lines.append(f"Name: {self.name_edit.text()}")
        lines.append(f"Gender: {self.gender_combo.currentText()}")
        lines.append(f"Date of Birth: {self.dob_edit.date().toString('dd/MM/yyyy')}")
        lines.append(f"Age: {self.age_label.text()}")
        lines.append(f"Section: {self.section_combo.currentText()}")
        lines.append(f"Admission Date: {self.admission_date_edit.date().toString('dd/MM/yyyy')}")
        lines.append(f"Current Location: {self.location_edit.text()}")
        lines.append(f"Report By: {self.report_by_edit.text()}")
        lines.append(f"Date Seen: {self.date_seen_edit.date().toString('dd/MM/yyyy')}")
        self.sent.emit("\n".join(lines))

    def set_clinician_details(self, details: dict):
        """Pre-fill clinician details."""
        if details.get("full_name"):
            self.report_by_edit.setText(details["full_name"])

    def fill_patient_info(self, patient_info: dict):
        """Fill fields from extracted patient demographics - only if fields are empty."""
        from datetime import datetime

        print(f"[GPRPatientDetailsPopup] fill_patient_info called with: {list(patient_info.keys())}")

        # Fill name if empty
        if patient_info.get("name") and not self.name_edit.text().strip():
            self.name_edit.setText(patient_info["name"])
            print(f"[GPRPatientDetailsPopup] Set name: {patient_info['name']}")

        # Fill gender if empty
        if patient_info.get("gender") and not self.gender_combo.currentText().strip():
            gender = patient_info["gender"].strip().title()
            index = self.gender_combo.findText(gender, Qt.MatchFlag.MatchContains)
            if index >= 0:
                self.gender_combo.setCurrentIndex(index)
            print(f"[GPRPatientDetailsPopup] Set gender: {gender}")

        # Fill DOB if at default (01/01/1980)
        if patient_info.get("dob"):
            current = self.dob_edit.date()
            if current == QDate(1980, 1, 1):
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    self.dob_edit.setDate(QDate(dob.year, dob.month, dob.day))
                    print(f"[GPRPatientDetailsPopup] Set DOB: {dob.strftime('%d/%m/%Y')}")

        # Fill MHA section if empty (combo box)
        if patient_info.get("mha_section") and not self.section_combo.currentText().strip():
            mha_section = patient_info["mha_section"]
            # Try to match to existing items first
            section_text = f"Section {mha_section}" if not mha_section.lower().startswith("section") else mha_section
            index = self.section_combo.findText(section_text, Qt.MatchFlag.MatchContains)
            if index >= 0:
                self.section_combo.setCurrentIndex(index)
            else:
                # Set as custom text (editable combo)
                self.section_combo.setCurrentText(section_text)
            print(f"[GPRPatientDetailsPopup] Set MHA section: {section_text}")

        # Fill hospital/ward location if empty
        if not self.location_edit.text().strip():
            hospital = patient_info.get("hospital", "")
            ward = patient_info.get("ward", "")
            location_parts = []
            if hospital:
                location_parts.append(hospital)
            if ward:
                location_parts.append(ward)
            if location_parts:
                location_text = ", ".join(location_parts)
                self.location_edit.setText(location_text)
                print(f"[GPRPatientDetailsPopup] Set location: {location_text}")


# ================================================================
# REPORT BASED ON POPUP
# ================================================================

class GPRReportBasedOnPopup(QWidget):
    """Popup for selecting what the report is based on."""

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        title = QLabel("This Report Is Based On")
        title.setStyleSheet("font-size: 20px; font-weight: 700; color: #2563eb;")
        layout.addWidget(title)

        self.checkboxes = {}

        options = [
            ("medical_reports", "Medical reports from the above"),
            ("interviews_nursing", "Interviews with nursing staff"),
            ("interviews_patient", "Interviews with the patient"),
            ("previous_notes_current", "Previous notes from current placement"),
            ("previous_notes_other", "Previous notes from other placements"),
            ("psychology_reports", "Psychology reports"),
            ("social_work_reports", "Social work reports"),
            ("ot_reports", "Occupational therapy reports"),
        ]

        for key, label in options:
            cb = QCheckBox(label)
            cb.setStyleSheet("""
                QCheckBox {
                    font-size: 19px;
                    background: transparent;
                }
                QCheckBox::indicator {
                    width: 18px;
                    height: 18px;
                }
            """)
            # Auto-sync to card when checkbox state changes
            cb.stateChanged.connect(self._send_to_report)
            self.checkboxes[key] = cb
            layout.addWidget(cb)

        layout.addStretch()

    def _send_to_report(self):
        selected = []
        for key, cb in self.checkboxes.items():
            if cb.isChecked():
                selected.append(f"☒ {cb.text()}")
            else:
                selected.append(f"☐ {cb.text()}")
        self.sent.emit("\n".join(selected))

    def set_entries(self, entries: list):
        """Set checkboxes based on extracted data."""
        if not entries:
            return

        # Combine all extracted text
        text = ""
        for entry in entries:
            if isinstance(entry, dict) and entry.get("text"):
                text += " " + entry["text"]
            elif isinstance(entry, str):
                text += " " + entry

        text_lower = text.lower()

        # Map keywords to checkbox keys
        keyword_map = {
            "medical_reports": ["medical reports", "medical report"],
            "interviews_nursing": ["nursing staff", "interviews with nursing"],
            "interviews_patient": ["with the patient", "interview with the patient"],
            "previous_notes_current": ["current placement", "notes from current"],
            "previous_notes_other": ["other placements", "notes from other"],
            "psychology_reports": ["psychology report", "psychological report"],
            "social_work_reports": ["social work report", "social worker report"],
            "ot_reports": ["occupational therapy", "ot report"],
        }

        # Check boxes based on keywords found with ☒ nearby
        for key, keywords in keyword_map.items():
            for keyword in keywords:
                # Find keyword in text and check if ☒ is nearby
                idx = text_lower.find(keyword)
                if idx != -1:
                    # Look for ☒ within 50 chars before or after
                    start = max(0, idx - 50)
                    end = min(len(text), idx + len(keyword) + 50)
                    snippet = text[start:end]
                    if "☒" in snippet:
                        self.checkboxes[key].setChecked(True)
                        break


# ================================================================
# PSYCH HISTORY POPUP (with admissions table)
# ================================================================

class GPRPsychHistoryPopup(QWidget):
    """Popup for past psychiatric history with detected admissions, clerking notes, and extracted data.

    Sends directly to card on click - no preview.
    """

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._entries = []
        self._extracted_checkboxes = []
        self._clerking_checkboxes = []
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: DETECTED ADMISSIONS (table only)
        # ====================================================
        self.detected_section = CollapsibleSection("Detected Admissions", start_collapsed=True)
        self.detected_section.set_content_height(180)
        self.detected_section._min_height = 100
        self.detected_section._max_height = 400
        self.detected_section.set_header_style("""
            QFrame {
                background: rgba(37, 99, 235, 0.15);
                border: 1px solid rgba(37, 99, 235, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.detected_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #1d4ed8;
                background: transparent;
                border: none;
            }
        """)

        detected_content = QWidget()
        detected_content.setStyleSheet("""
            QWidget {
                background: rgba(239, 246, 255, 0.95);
                border: 1px solid rgba(37, 99, 235, 0.3);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        detected_layout = QVBoxLayout(detected_content)
        detected_layout.setContentsMargins(12, 12, 12, 12)
        detected_layout.setSpacing(8)

        # Detected admissions table (read-only, auto-populated)
        self.detected_table = QTableWidget(0, 3)
        self.detected_table.setHorizontalHeaderLabels(["Admission Date", "Discharge Date", "Duration"])
        self.detected_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.detected_table.setMinimumHeight(60)
        self.detected_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.detected_table.setStyleSheet("""
            QTableWidget {
                background: white;
                border: 1px solid #93c5fd;
                border-radius: 4px;
            }
            QHeaderView::section {
                background: #dbeafe;
                padding: 6px;
                border: none;
                border-bottom: 1px solid #93c5fd;
                font-weight: 600;
                color: #1e40af;
            }
        """)
        detected_layout.addWidget(self.detected_table)

        # Export to Preview container
        export_container = QFrame()
        export_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.6);
                border: 1px solid #93c5fd;
                border-radius: 6px;
            }
        """)
        export_layout = QHBoxLayout(export_container)
        export_layout.setContentsMargins(8, 4, 8, 4)

        self.export_table_cb = QCheckBox("Export to Preview")
        self.export_table_cb.setStyleSheet("""
            QCheckBox {
                font-size: 16px;
                font-weight: 600;
                color: #1e40af;
                background: transparent;
            }
            QCheckBox::indicator {
                width: 14px;
                height: 14px;
            }
        """)
        self.export_table_cb.stateChanged.connect(self._update_preview)
        export_layout.addWidget(self.export_table_cb)
        export_layout.addStretch()

        detected_layout.addWidget(export_container)

        self.detected_section.set_content(detected_content)
        self.detected_section.setVisible(False)  # Hidden until data loaded
        main_layout.addWidget(self.detected_section)

        # Add spacing between detected admissions and clerking notes
        spacer1 = QWidget()
        spacer1.setFixedHeight(4)
        spacer1.setStyleSheet("background: transparent;")
        main_layout.addWidget(spacer1)

        # ====================================================
        # SECTION 2B: ADMISSION CLERKING NOTES (separate section)
        # ====================================================
        self.clerking_section = CollapsibleSection("Admission Clerking Notes", start_collapsed=True)
        self.clerking_section.set_content_height(250)
        self.clerking_section._min_height = 100
        self.clerking_section._max_height = 500
        self.clerking_section.set_header_style("""
            QFrame {
                background: rgba(37, 99, 235, 0.15);
                border: 1px solid rgba(37, 99, 235, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.clerking_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #1d4ed8;
                background: transparent;
                border: none;
            }
        """)

        clerking_content = QWidget()
        clerking_content.setStyleSheet("""
            QWidget {
                background: rgba(239, 246, 255, 0.95);
                border: 1px solid rgba(37, 99, 235, 0.3);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        clerking_content_layout = QVBoxLayout(clerking_content)
        clerking_content_layout.setContentsMargins(12, 12, 12, 12)
        clerking_content_layout.setSpacing(8)

        # Scrollable container for clerking entries
        clerking_scroll = QScrollArea()
        clerking_scroll.setWidgetResizable(True)
        clerking_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        clerking_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        clerking_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        clerking_scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollArea > QWidget > QWidget { background: transparent; }
        """)

        self.clerking_container = QWidget()
        self.clerking_container.setStyleSheet("background: transparent;")
        self.clerking_entries_layout = QVBoxLayout(self.clerking_container)
        self.clerking_entries_layout.setContentsMargins(2, 2, 2, 2)
        self.clerking_entries_layout.setSpacing(8)
        self.clerking_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        clerking_scroll.setWidget(self.clerking_container)
        clerking_content_layout.addWidget(clerking_scroll)

        self.clerking_section.set_content(clerking_content)
        self.clerking_section.setVisible(False)  # Hidden until data loaded
        main_layout.addWidget(self.clerking_section)

        # ====================================================
        # SECTION 3: IMPORTED DATA
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(150)
        self.extracted_section._min_height = 60
        self.extracted_section._max_height = 400
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        extracted_content = QWidget()
        extracted_content.setStyleSheet("""
            QWidget {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
            QCheckBox {
                background: transparent;
                border: none;
                padding: 4px;
                font-size: 17px;
                color: #4a4a4a;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
            }
        """)

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollArea > QWidget > QWidget { background: transparent; }
        """)

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)
        self.extracted_checkboxes_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        table_parts = []
        clerking_parts = []
        extracted_parts = []

        # Export table if checkbox is checked
        if hasattr(self, 'export_table_cb') and self.export_table_cb.isChecked():
            table_lines = []
            for row in range(self.detected_table.rowCount()):
                adm_date = self.detected_table.item(row, 0)
                dis_date = self.detected_table.item(row, 1)
                duration = self.detected_table.item(row, 2)

                adm_str = adm_date.text() if adm_date else ""
                dis_str = dis_date.text() if dis_date else ""
                dur_str = duration.text() if duration else ""

                if adm_str:
                    table_lines.append(f"Admission {row + 1}: {adm_str} - {dis_str} ({dur_str})")

            if table_lines:
                table_parts.append("HOSPITAL ADMISSIONS:\n" + "\n".join(table_lines))

        # Checked clerking entries
        for cb in self._clerking_checkboxes:
            if cb.isChecked():
                clerking_parts.append(cb.property("full_text"))

        # Checked extracted entries
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                extracted_parts.append(cb.property("full_text"))

        # Combine
        all_parts = []
        if table_parts:
            all_parts.extend(table_parts)
        if clerking_parts:
            all_parts.append("ADMISSION NOTES:\n" + "\n\n".join(clerking_parts))
        if extracted_parts:
            all_parts.append("FROM NOTES:\n" + "\n\n".join(extracted_parts))

        combined = "\n\n".join(all_parts) if all_parts else ""
        # Send directly to card
        self.sent.emit(combined)

    def set_notes(self, notes: list):
        """Analyze notes using timeline to detect admissions and find clerking notes."""
        from timeline_builder import build_rio_timeline
        from datetime import timedelta

        print(f"[GPR-PSYCH] set_notes called with {len(notes) if notes else 0} notes")

        if not notes:
            self.detected_section.setVisible(False)
            return

        # Run timeline analysis
        try:
            episodes = build_rio_timeline(notes, debug=False)
            print(f"[GPR-PSYCH] Timeline returned {len(episodes)} episodes")
        except Exception as e:
            print(f"[GPR-PSYCH] Timeline error: {e}")
            self.detected_section.setVisible(False)
            return

        # Filter for inpatient admissions only
        admissions = [ep for ep in episodes if ep.get("type") == "inpatient"]
        print(f"[GPR-PSYCH] Found {len(admissions)} inpatient admissions")

        if not admissions:
            self.detected_section.setVisible(False)
            return

        # Show the section
        self.detected_section.setVisible(True)
        print(f"[GPR-PSYCH] Detected section now visible")

        # Clear and populate detected admissions table
        self.detected_table.setRowCount(len(admissions))

        for row, adm in enumerate(admissions):
            start_date = adm.get("start")
            end_date = adm.get("end")

            if start_date:
                start_str = start_date.strftime("%d %b %Y") if hasattr(start_date, "strftime") else str(start_date)
            else:
                start_str = "Unknown"

            if end_date:
                end_str = end_date.strftime("%d %b %Y") if hasattr(end_date, "strftime") else str(end_date)
            else:
                end_str = "Ongoing"

            if start_date and end_date:
                try:
                    duration_days = (end_date - start_date).days
                    if duration_days < 7:
                        duration_str = f"{duration_days} days"
                    elif duration_days < 30:
                        weeks = duration_days // 7
                        duration_str = f"{weeks} week{'s' if weeks > 1 else ''}"
                    else:
                        months = duration_days // 30
                        duration_str = f"{months} month{'s' if months > 1 else ''}"
                except:
                    duration_str = "Unknown"
            else:
                duration_str = "Ongoing"

            self.detected_table.setItem(row, 0, QTableWidgetItem(start_str))
            self.detected_table.setItem(row, 1, QTableWidgetItem(end_str))
            self.detected_table.setItem(row, 2, QTableWidgetItem(duration_str))

        # Find clerking/admission notes for each admission
        for cb in self._clerking_checkboxes:
            cb.deleteLater()
        self._clerking_checkboxes.clear()

        while self.clerking_entries_layout.count():
            item = self.clerking_entries_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Admission entry keywords
        ADMISSION_KEYWORDS = [
            "admission to ward", "admitted to ward", "admitted to the ward",
            "brought to ward", "brought to the ward", "brought into ward",
            "brought onto ward", "brought onto the ward",
            "arrived on ward", "arrived on the ward", "arrived to ward",
            "transferred to ward", "transferred to the ward",
            "escorted to ward", "escorted to the ward",
            "on admission", "admission clerking", "clerking",
            "duty doctor admission", "admission note",
            "accepted to ward", "accepted onto ward",
            "admitted under", "accepted under",
            "detained under", "sectioned", "section 2", "section 3",
            "136 suite", "sec 136", "section 136",
            "nursing admission", "admission assessment",
            "initial assessment", "ward admission",
            "new admission", "patient admitted",
        ]

        clerking_notes = []
        seen_keys = set()

        for adm in admissions:
            adm_start = adm.get("start")
            if not adm_start:
                continue

            window_end = adm_start + timedelta(days=14)

            admission_window_notes = []
            for note in notes:
                note_date = note.get("date")
                if not note_date:
                    continue

                if hasattr(note_date, "date"):
                    note_date_obj = note_date.date()
                else:
                    note_date_obj = note_date

                if adm_start <= note_date_obj <= window_end:
                    admission_window_notes.append((note_date_obj, note))

            admission_window_notes.sort(key=lambda x: x[0])

            found_admission_note = None
            for note_date_obj, note in admission_window_notes:
                text = (note.get("text", "") or note.get("content", "")).lower()

                if any(kw in text for kw in ADMISSION_KEYWORDS):
                    key = (note_date_obj, text[:100])
                    if key not in seen_keys:
                        seen_keys.add(key)
                        found_admission_note = {
                            "date": note.get("date"),
                            "text": note.get("text", "") or note.get("content", ""),
                            "admission_label": adm.get("label", "Admission")
                        }
                    break

            if found_admission_note:
                clerking_notes.append(found_admission_note)

        # Create collapsible entry boxes for each clerking note (blue UI)
        for clerking in clerking_notes:
            dt = clerking.get("date")
            text = clerking.get("text", "").strip()
            adm_label = clerking.get("admission_label", "")

            if not text:
                continue

            if dt:
                if hasattr(dt, "strftime"):
                    date_str = dt.strftime("%d %b %Y")
                else:
                    date_str = str(dt)
            else:
                date_str = "No date"

            entry_frame = QFrame()
            entry_frame.setObjectName("clerkingEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet("""
                QFrame#clerkingEntryFrame {
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(37, 99, 235, 0.4);
                    border-radius: 8px;
                    padding: 4px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button on the LEFT
            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(37, 99, 235, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 18px;
                    font-weight: bold;
                    color: #1e40af;
                }
                QPushButton:hover { background: rgba(37, 99, 235, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            if adm_label:
                badge = QLabel(adm_label)
                badge.setStyleSheet("""
                    QLabel {
                        font-size: 14px;
                        font-weight: 600;
                        color: white;
                        background: #2563eb;
                        border: none;
                        border-radius: 3px;
                        padding: 2px 6px;
                    }
                """)
                header_row.addWidget(badge)

            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 17px;
                    font-weight: 600;
                    color: #1e40af;
                    background: transparent;
                    border: none;
                }
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)
            header_row.addStretch()

            # Checkbox on the RIGHT
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(self._update_preview)
            header_row.addWidget(cb)

            entry_layout.addLayout(header_row)

            body_text = QTextEdit()
            body_text.setPlainText(text)
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet("""
                QTextEdit {
                    font-size: 17px;
                    color: #333;
                    background: rgba(219, 234, 254, 0.5);
                    border: none;
                    padding: 8px;
                    border-radius: 6px;
                }
            """)
            body_text.document().setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            doc_height = body_text.document().size().height() + 20
            body_text.setFixedHeight(int(max(doc_height, 60)))
            body_text.setVisible(False)
            entry_layout.addWidget(body_text)

            drag_bar = QFrame()
            drag_bar.setFixedHeight(8)
            drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
            drag_bar.setStyleSheet("""
                QFrame {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(37,99,235,0.1), stop:0.5 rgba(37,99,235,0.3), stop:1 rgba(37,99,235,0.1));
                    border-radius: 2px; margin: 2px 40px;
                }
                QFrame:hover {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(37,99,235,0.2), stop:0.5 rgba(37,99,235,0.5), stop:1 rgba(37,99,235,0.2));
                }
            """)
            drag_bar.setVisible(False)
            drag_bar._drag_y = None
            drag_bar._init_h = None
            def _make_drag_handlers(handle, text_widget):
                def press(ev):
                    handle._drag_y = ev.globalPosition().y()
                    handle._init_h = text_widget.height()
                def move(ev):
                    if handle._drag_y is not None:
                        delta = int(ev.globalPosition().y() - handle._drag_y)
                        new_h = max(60, handle._init_h + delta)
                        text_widget.setMinimumHeight(new_h)
                        text_widget.setMaximumHeight(new_h)
                def release(ev):
                    if handle._drag_y is not None:
                        text_widget.setMaximumHeight(16777215)
                        handle._drag_y = None
                return press, move, release
            dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
            drag_bar.mousePressEvent = dp
            drag_bar.mouseMoveEvent = dm
            drag_bar.mouseReleaseEvent = dr
            entry_layout.addWidget(drag_bar)

            def make_toggle(btn, body, frame, popup_self, bar):
                def toggle():
                    is_visible = body.isVisible()
                    body.setVisible(not is_visible)
                    bar.setVisible(not is_visible)
                    btn.setText("▾" if not is_visible else "▸")
                    frame.updateGeometry()
                    if hasattr(popup_self, 'clerking_container'):
                        popup_self.clerking_container.updateGeometry()
                        popup_self.clerking_container.update()
                return toggle

            toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
            toggle_btn.clicked.connect(toggle_fn)
            date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

            self.clerking_entries_layout.addWidget(entry_frame)
            self._clerking_checkboxes.append(cb)

        # Show the sections (both stay collapsed)
        self.detected_section.setVisible(True)
        if clerking_notes:
            self.clerking_section.setVisible(True)

        print(f"[GPR-PSYCH] Detected {len(admissions)} admissions, {len(clerking_notes)} clerking notes")

    def set_entries(self, items: list):
        """Display extracted data from notes with collapsible dated entry boxes."""
        self._entries = items

        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if isinstance(items, str):
            if items.strip():
                items = [{"date": None, "text": p.strip()} for p in items.split("\n\n") if p.strip()]
            else:
                items = []

        if items:
            def get_sort_date(item):
                dt = item.get("date")
                if dt is None:
                    return ""
                if hasattr(dt, "strftime"):
                    return dt.strftime("%Y-%m-%d")
                return str(dt)

            sorted_items = sorted(items, key=get_sort_date, reverse=True)

            for item in sorted_items:
                dt = item.get("date")
                text = item.get("text", "").strip()
                if not text:
                    continue

                if dt:
                    if hasattr(dt, "strftime"):
                        date_str = dt.strftime("%d %b %Y")
                    else:
                        date_str = str(dt)
                else:
                    date_str = "No date"

                entry_frame = QFrame()
                entry_frame.setObjectName("entryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
                entry_frame.setStyleSheet("""
                    QFrame#entryFrame {
                        background: rgba(255, 255, 255, 0.95);
                        border: 1px solid rgba(180, 150, 50, 0.4);
                        border-radius: 8px;
                        padding: 4px;
                    }
                """)
                entry_layout = QVBoxLayout(entry_frame)
                entry_layout.setContentsMargins(10, 8, 10, 8)
                entry_layout.setSpacing(6)
                entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(180, 150, 50, 0.2);
                        border: none;
                        border-radius: 4px;
                        font-size: 17px;
                        font-weight: bold;
                        color: #806000;
                    }
                    QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
                """)
                header_row.addWidget(toggle_btn)

                date_label = QLabel(f"📅 {date_str}")
                date_label.setStyleSheet("""
                    QLabel {
                        font-size: 17px;
                        font-weight: 600;
                        color: #806000;
                        background: transparent;
                        border: none;
                    }
                """)
                date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                header_row.addWidget(date_label)
                header_row.addStretch()

                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("""
                    QCheckBox { background: transparent; }
                    QCheckBox::indicator { width: 16px; height: 16px; }
                """)
                cb.stateChanged.connect(self._update_preview)
                header_row.addWidget(cb)

                entry_layout.addLayout(header_row)

                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setStyleSheet("""
                    QTextEdit {
                        font-size: 17px;
                        color: #333;
                        background: rgba(255, 248, 220, 0.5);
                        border: none;
                        padding: 8px;
                        border-radius: 6px;
                    }
                """)
                body_text.document().setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
                doc_height = body_text.document().size().height() + 20
                body_text.setFixedHeight(int(max(doc_height, 60)))
                body_text.setVisible(False)
                entry_layout.addWidget(body_text)

                drag_bar = QFrame()
                drag_bar.setFixedHeight(8)
                drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
                drag_bar.setStyleSheet("""
                    QFrame {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                        border-radius: 2px; margin: 2px 40px;
                    }
                    QFrame:hover {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                    }
                """)
                drag_bar.setVisible(False)
                drag_bar._drag_y = None
                drag_bar._init_h = None
                def _make_drag_handlers(handle, text_widget):
                    def press(ev):
                        handle._drag_y = ev.globalPosition().y()
                        handle._init_h = text_widget.height()
                    def move(ev):
                        if handle._drag_y is not None:
                            delta = int(ev.globalPosition().y() - handle._drag_y)
                            new_h = max(60, handle._init_h + delta)
                            text_widget.setMinimumHeight(new_h)
                            text_widget.setMaximumHeight(new_h)
                    def release(ev):
                        if handle._drag_y is not None:
                            text_widget.setMaximumHeight(16777215)
                            handle._drag_y = None
                    return press, move, release
                dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
                drag_bar.mousePressEvent = dp
                drag_bar.mouseMoveEvent = dm
                drag_bar.mouseReleaseEvent = dr
                entry_layout.addWidget(drag_bar)

                def make_toggle(btn, body, frame, popup_self, bar):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        bar.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                        if hasattr(popup_self, 'extracted_container'):
                            popup_self.extracted_container.updateGeometry()
                            popup_self.extracted_container.update()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
                toggle_btn.clicked.connect(toggle_fn)
                date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                self.extracted_checkboxes_layout.addWidget(entry_frame)
                self._extracted_checkboxes.append(cb)

            self.extracted_section.setVisible(True)
            # Keep collapsed on open
            # if self.extracted_section._is_collapsed:
            #     self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)


# ================================================================
# GPR RISK POPUP (Current and Historical risk)
# ================================================================

class GPRRiskPopup(QWidget):
    """Popup for Risk section with Current and Historical risk factors - follows DischargeRiskPopup pattern."""

    sent = Signal(str)

    RISK_TYPES = [
        ("violence", "Violence to others"),
        ("verbal_aggression", "Verbal aggression"),
        ("self_harm", "Self-harm"),
        ("suicide", "Suicide"),
        ("self_neglect", "Self-neglect"),
        ("exploitation", "Exploitation by others"),
        ("sexual", "Sexually inappropriate behaviour"),
        ("substance", "Substance misuse"),
        ("property_damage", "Property damage"),
        ("awol", "AWOL/Absconding"),
        ("deterioration", "Mental health deterioration"),
        ("non_compliance", "Non-compliance with treatment"),
    ]

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._entries = []
        self._extracted_checkboxes = []
        self._current_widgets = {}
        self._historical_widgets = {}
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        from PySide6.QtWidgets import QSlider

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area for all sections
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        self.main_layout = QVBoxLayout(main_container)
        self.main_layout.setContentsMargins(4, 4, 4, 4)
        self.main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: CURRENT RISK (collapsible with drag bar)
        # ====================================================
        self.current_section = CollapsibleSection("Current Risk", start_collapsed=True)
        self.current_section.set_content_height(200)
        self.current_section._min_height = 80
        self.current_section._max_height = 400
        self.current_section.set_header_style("""
            QFrame {
                background: rgba(220, 38, 38, 0.15);
                border: 1px solid rgba(220, 38, 38, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.current_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #991b1b;
                background: transparent;
                border: none;
            }
        """)

        current_content = QWidget()
        current_content.setStyleSheet("""
            QWidget {
                background: rgba(255,255,255,0.92);
                border: 1px solid rgba(0,0,0,0.15);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        current_outer_layout = QVBoxLayout(current_content)
        current_outer_layout.setContentsMargins(0, 0, 0, 0)
        current_outer_layout.setSpacing(0)

        current_scroll = QScrollArea()
        current_scroll.setWidgetResizable(True)
        current_scroll.setFrameShape(QScrollArea.NoFrame)
        current_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        current_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        current_scroll.setStyleSheet("background: transparent; border: none;")

        current_inner = QWidget()
        current_inner.setStyleSheet("background: transparent;")
        current_layout = QVBoxLayout(current_inner)
        current_layout.setContentsMargins(12, 12, 12, 12)
        current_layout.setSpacing(6)

        for key, label in self.RISK_TYPES:
            container = QWidget()
            container.setStyleSheet("background: transparent;")
            container_layout = QVBoxLayout(container)
            container_layout.setContentsMargins(0, 0, 0, 0)
            container_layout.setSpacing(4)

            cb = QCheckBox(label)
            cb.setStyleSheet("""
                QCheckBox {
                    font-size: 19px;
                    background: transparent;
                }
                QCheckBox::indicator {
                    width: 18px;
                    height: 18px;
                }
            """)
            cb.toggled.connect(self._update_preview)
            container_layout.addWidget(cb)

            # Severity slider (hidden by default)
            slider_container = QWidget()
            slider_container.setStyleSheet("background: transparent;")
            slider_outer = QVBoxLayout(slider_container)
            slider_outer.setContentsMargins(20, 0, 0, 8)
            slider_outer.setSpacing(2)

            # Top row: label + slider
            slider_row = QHBoxLayout()
            slider_row.setSpacing(8)

            slider_lbl = QLabel("Severity:")
            slider_lbl.setStyleSheet("font-size: 17px; color: #6b7280; background: transparent;")
            slider_row.addWidget(slider_lbl)

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setMinimum(1)
            slider.setMaximum(3)
            slider.setValue(2)
            slider.setTickPosition(QSlider.TickPosition.TicksBelow)
            slider.setTickInterval(1)
            slider.setFixedWidth(120)
            slider.valueChanged.connect(self._update_preview)
            # Sync current slider with historical slider
            slider.valueChanged.connect(lambda v, k=key: self._sync_historical_slider(k, v))
            slider_row.addWidget(slider)
            slider_row.addStretch()

            slider_outer.addLayout(slider_row)

            # Bottom row: level label (below slider)
            level_lbl = QLabel("Medium")
            level_lbl.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500; margin-left: 52px; background: transparent;")
            slider.valueChanged.connect(lambda v, l=level_lbl: l.setText(["Low", "Medium", "High"][v-1]))
            slider_outer.addWidget(level_lbl)

            slider_container.hide()
            container_layout.addWidget(slider_container)

            cb.toggled.connect(lambda checked, sc=slider_container: sc.setVisible(checked))
            # When current risk is checked, also check historical at same severity
            cb.toggled.connect(lambda checked, k=key: self._on_current_risk_toggled(k, checked))

            current_layout.addWidget(container)

            self._current_widgets[key] = {
                "checkbox": cb,
                "slider": slider,
                "slider_container": slider_container,
                "level_label": level_lbl
            }

        current_layout.addStretch()
        current_scroll.setWidget(current_inner)
        current_outer_layout.addWidget(current_scroll)

        self.current_section.set_content(current_content)
        self.main_layout.addWidget(self.current_section)

        # ====================================================
        # SECTION 3: HISTORICAL RISK (collapsible with drag bar)
        # ====================================================
        self.historical_section = CollapsibleSection("Historical Risk", start_collapsed=True)
        self.historical_section.set_content_height(200)
        self.historical_section._min_height = 80
        self.historical_section._max_height = 400
        self.historical_section.set_header_style("""
            QFrame {
                background: rgba(245, 158, 11, 0.15);
                border: 1px solid rgba(245, 158, 11, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.historical_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #92400e;
                background: transparent;
                border: none;
            }
        """)

        historical_content = QWidget()
        historical_content.setStyleSheet("""
            QWidget {
                background: rgba(255,255,255,0.92);
                border: 1px solid rgba(0,0,0,0.15);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        historical_outer_layout = QVBoxLayout(historical_content)
        historical_outer_layout.setContentsMargins(0, 0, 0, 0)
        historical_outer_layout.setSpacing(0)

        historical_scroll = QScrollArea()
        historical_scroll.setWidgetResizable(True)
        historical_scroll.setFrameShape(QScrollArea.NoFrame)
        historical_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        historical_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        historical_scroll.setStyleSheet("background: transparent; border: none;")

        historical_inner = QWidget()
        historical_inner.setStyleSheet("background: transparent;")
        historical_layout = QVBoxLayout(historical_inner)
        historical_layout.setContentsMargins(12, 12, 12, 12)
        historical_layout.setSpacing(6)

        for key, label in self.RISK_TYPES:
            container = QWidget()
            container.setStyleSheet("background: transparent;")
            container_layout = QVBoxLayout(container)
            container_layout.setContentsMargins(0, 0, 0, 0)
            container_layout.setSpacing(4)

            cb = QCheckBox(label)
            cb.setStyleSheet("""
                QCheckBox {
                    font-size: 19px;
                    background: transparent;
                }
                QCheckBox::indicator {
                    width: 18px;
                    height: 18px;
                }
            """)
            cb.toggled.connect(self._update_preview)
            container_layout.addWidget(cb)

            # Severity slider (hidden by default)
            slider_container = QWidget()
            slider_container.setStyleSheet("background: transparent;")
            slider_outer = QVBoxLayout(slider_container)
            slider_outer.setContentsMargins(20, 0, 0, 8)
            slider_outer.setSpacing(2)

            # Top row: label + slider
            slider_row = QHBoxLayout()
            slider_row.setSpacing(8)

            slider_lbl = QLabel("Severity:")
            slider_lbl.setStyleSheet("font-size: 17px; color: #6b7280; background: transparent;")
            slider_row.addWidget(slider_lbl)

            slider = NoWheelSlider(Qt.Orientation.Horizontal)
            slider.setMinimum(1)
            slider.setMaximum(3)
            slider.setValue(2)
            slider.setTickPosition(QSlider.TickPosition.TicksBelow)
            slider.setTickInterval(1)
            slider.setFixedWidth(120)
            slider.valueChanged.connect(self._update_preview)
            slider_row.addWidget(slider)
            slider_row.addStretch()

            slider_outer.addLayout(slider_row)

            # Bottom row: level label (below slider)
            level_lbl = QLabel("Medium")
            level_lbl.setStyleSheet("font-size: 17px; color: #374151; font-weight: 500; margin-left: 52px; background: transparent;")
            slider.valueChanged.connect(lambda v, l=level_lbl: l.setText(["Low", "Medium", "High"][v-1]))
            slider_outer.addWidget(level_lbl)

            slider_container.hide()
            container_layout.addWidget(slider_container)

            cb.toggled.connect(lambda checked, sc=slider_container: sc.setVisible(checked))

            historical_layout.addWidget(container)

            self._historical_widgets[key] = {
                "checkbox": cb,
                "slider": slider,
                "slider_container": slider_container,
                "level_label": level_lbl
            }

        historical_layout.addStretch()
        historical_scroll.setWidget(historical_inner)
        historical_outer_layout.addWidget(historical_scroll)

        self.historical_section.set_content(historical_content)
        self.main_layout.addWidget(self.historical_section)

        # ====================================================
        # SECTION 4: IMPORTED DATA (collapsible with drag bar)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(250)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 500
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        # Use QScrollArea directly as content
        self.extracted_scroll = QScrollArea()
        self.extracted_scroll.setWidgetResizable(True)
        self.extracted_scroll.setFrameShape(QScrollArea.NoFrame)
        self.extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.extracted_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
            QCheckBox {
                background: transparent;
                border: none;
                padding: 4px;
                font-size: 18px;
                color: #4a4a4a;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
            }
        """)

        # Container for checkboxes inside scroll area
        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(6, 6, 6, 6)
        self.extracted_checkboxes_layout.setSpacing(8)

        self.extracted_scroll.setWidget(self.extracted_container)
        self.extracted_section.set_content(self.extracted_scroll)
        self.extracted_section.setVisible(False)
        self.main_layout.addWidget(self.extracted_section)

        # Finalize main scroll area
        self.main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _build_risk_narrative(self, risks: list, is_historical: bool = False) -> str:
        """Build a narrative sentence from a list of (risk_name, severity_value) tuples."""
        if not risks:
            return ""

        # Sort by severity (high=3 first, then moderate=2, then low=1)
        sorted_risks = sorted(risks, key=lambda x: x[1], reverse=True)

        severity_words = {1: "low", 2: "moderate", 3: "high"}

        # Group by severity
        high_risks = [r[0] for r in sorted_risks if r[1] == 3]
        moderate_risks = [r[0] for r in sorted_risks if r[1] == 2]
        low_risks = [r[0] for r in sorted_risks if r[1] == 1]

        parts = []
        prefix = "Historically, the" if is_historical else "The"

        def join_risks(risk_list):
            if len(risk_list) == 1:
                return risk_list[0]
            elif len(risk_list) == 2:
                return f"{risk_list[0]} and {risk_list[1]}"
            else:
                return ", ".join(risk_list[:-1]) + f", and {risk_list[-1]}"

        if high_risks:
            if len(high_risks) == 1:
                parts.append(f"risk of {high_risks[0]} is high")
            else:
                parts.append(f"risks of {join_risks(high_risks)} are high")

        if moderate_risks:
            if len(moderate_risks) == 1:
                if parts:
                    parts.append(f"{moderate_risks[0]} is moderate")
                else:
                    parts.append(f"risk of {moderate_risks[0]} is moderate")
            else:
                if parts:
                    parts.append(f"{join_risks(moderate_risks)} are moderate")
                else:
                    parts.append(f"risks of {join_risks(moderate_risks)} are moderate")

        if low_risks:
            if len(low_risks) == 1:
                if parts:
                    parts.append(f"{low_risks[0]} is low")
                else:
                    parts.append(f"risk of {low_risks[0]} is low")
            else:
                if parts:
                    parts.append(f"{join_risks(low_risks)} are low")
                else:
                    parts.append(f"risks of {join_risks(low_risks)} are low")

        if len(parts) == 1:
            return f"{prefix} {parts[0]}."
        elif len(parts) == 2:
            return f"{prefix} {parts[0]}, and {parts[1]}."
        else:
            return f"{prefix} {parts[0]}, {parts[1]}, and {parts[2]}."

    def _on_current_risk_toggled(self, key: str, checked: bool):
        """When a current risk is checked, also check the corresponding historical risk at same severity."""
        if key not in self._historical_widgets:
            return

        historical = self._historical_widgets[key]
        if checked:
            # Get current severity level
            current_slider = self._current_widgets[key]["slider"]
            severity = current_slider.value()

            # Check historical checkbox and set same severity
            historical["checkbox"].blockSignals(True)
            historical["checkbox"].setChecked(True)
            historical["checkbox"].blockSignals(False)

            # Show slider container
            historical["slider_container"].setVisible(True)

            # Set slider to same severity
            historical["slider"].blockSignals(True)
            historical["slider"].setValue(severity)
            historical["slider"].blockSignals(False)

            # Update level label
            historical["level_label"].setText(["Low", "Medium", "High"][severity - 1])

            # Trigger preview update
            self._update_preview()

    def _sync_historical_slider(self, key: str, value: int):
        """When current risk slider changes, sync historical slider to same value."""
        if key not in self._historical_widgets:
            return

        # Only sync if current checkbox is checked
        if not self._current_widgets[key]["checkbox"].isChecked():
            return

        historical = self._historical_widgets[key]

        # Only sync if historical is also checked
        if historical["checkbox"].isChecked():
            historical["slider"].blockSignals(True)
            historical["slider"].setValue(value)
            historical["slider"].blockSignals(False)
            historical["level_label"].setText(["Low", "Medium", "High"][value - 1])

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        current_risks = []
        historical_risks = []
        extracted_parts = []

        # Current risks - collect (risk_name, severity_value)
        for key, widgets in self._current_widgets.items():
            if widgets["checkbox"].isChecked():
                risk_name = widgets["checkbox"].text().lower()
                severity_val = widgets["slider"].value()
                current_risks.append((risk_name, severity_val))

        # Historical risks - collect (risk_name, severity_value)
        for key, widgets in self._historical_widgets.items():
            if widgets["checkbox"].isChecked():
                risk_name = widgets["checkbox"].text().lower()
                severity_val = widgets["slider"].value()
                historical_risks.append((risk_name, severity_val))

        # Checked extracted entries
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                extracted_parts.append(cb.property("full_text"))

        # Build combined text
        sections = []
        if current_risks:
            narrative = self._build_risk_narrative(current_risks, is_historical=False)
            sections.append(f"CURRENT RISK:\n{narrative}")
        if historical_risks:
            narrative = self._build_risk_narrative(historical_risks, is_historical=True)
            sections.append(f"HISTORICAL RISK:\n{narrative}")
        if extracted_parts:
            sections.append("\n".join(extracted_parts))

        combined = "\n\n".join(sections) if sections else ""
        # Send directly to card
        self.sent.emit(combined)

    def set_entries(self, entries: list):
        """Set entries - but for Risk we prefer to use set_notes_for_risk_analysis."""
        self._entries = entries
        # For backwards compatibility, just store - actual display is via set_notes_for_risk_analysis

    def set_notes_for_risk_analysis(self, notes: list):
        """Run risk analysis on notes and display results in date order with risk type and highlighted matches.

        Also auto-populates Current Risk (last 3 months) and Historical Risk (all time) checkboxes.
        """
        from risk_overview_panel import analyze_notes_for_risk
        from datetime import timedelta
        import re
        import html

        # Clear existing checkboxes
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        # Clear layout
        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not notes:
            self.extracted_section.setVisible(False)
            return

        # Run risk analysis
        results = analyze_notes_for_risk(notes)

        # Mapping from risk_overview categories/subcategories to GPRRiskPopup RISK_TYPES keys
        # Some categories map to multiple risk types based on subcategory
        # Mapping from risk_overview categories to GPRRiskPopup RISK_TYPES keys
        # Category names must match EXACTLY with risk_overview_panel.py RISK_CATEGORIES
        CATEGORY_TO_RISK_TYPE = {
            "Physical Aggression": "violence",
            "Verbal Aggression": "verbal_aggression",
            "Self-Harm": "self_harm",
            "Sexual Behaviour": "sexual",
            "Self-Neglect": "self_neglect",
            "AWOL/Absconding": "awol",
            "Property Damage": "property_damage",
            "Substance Misuse": "substance",
            "Non-Compliance": "non_compliance",
            "Bullying/Exploitation": "exploitation",
        }

        # Subcategory overrides - some subcategories map to different risk types
        SUBCATEGORY_OVERRIDES = {
            "Ligature": "suicide",  # Self-harm ligature → suicide
            "Overdose": "suicide",  # Overdose → suicide
            "Suicide Attempt": "suicide",
            "Suicidal Ideation": "suicide",
        }

        # Collect all incidents and count by risk type
        all_incidents = []
        from collections import defaultdict
        current_risk_counts = defaultdict(int)  # Count in last 3 months
        historical_risk_counts = defaultdict(int)  # Count ALL time

        # Find the latest and earliest dates
        latest_date = None
        earliest_date = None
        for cat_name, cat_data in results.get("categories", {}).items():
            for incident in cat_data.get("incidents", []):
                inc_date = incident.get("date")
                if inc_date:
                    if latest_date is None or inc_date > latest_date:
                        latest_date = inc_date
                    if earliest_date is None or inc_date < earliest_date:
                        earliest_date = inc_date

        # Calculate 3-month cutoff from latest entry
        three_months_ago = None
        if latest_date:
            three_months_ago = latest_date - timedelta(days=90)

        # Calculate number of years for historical averaging
        num_years = 1  # Default to 1 year minimum
        if latest_date and earliest_date:
            days_span = (latest_date - earliest_date).days
            num_years = max(1, days_span / 365.0)  # At least 1 year

        for cat_name, cat_data in results.get("categories", {}).items():
            for incident in cat_data.get("incidents", []):
                inc_date = incident.get("date")
                subcat = incident.get("subcategory", "")

                # Determine risk type - check subcategory overrides first
                risk_type_key = None
                for override_key, override_value in SUBCATEGORY_OVERRIDES.items():
                    if override_key.lower() in subcat.lower():
                        risk_type_key = override_value
                        break

                # Fall back to category mapping
                if not risk_type_key:
                    risk_type_key = CATEGORY_TO_RISK_TYPE.get(cat_name)

                if risk_type_key:
                    # Historical counts ALL entries
                    historical_risk_counts[risk_type_key] += 1

                    # Current counts only last 3 months
                    if inc_date and three_months_ago and inc_date >= three_months_ago:
                        current_risk_counts[risk_type_key] += 1

                all_incidents.append({
                    "date": inc_date,
                    "text": incident.get("full_text", ""),
                    "matched": incident.get("matched", ""),
                    "subcategory": subcat,
                    "severity": incident.get("severity", "medium"),
                    "category": cat_name,
                })

        # Auto-populate Current Risk checkboxes (last 3 months)
        # Severity: < 3 = MEDIUM, >= 3 = HIGH, AWOL 1+ = HIGH
        for risk_key, count in current_risk_counts.items():
            if risk_key in self._current_widgets and count > 0:
                self._current_widgets[risk_key]["checkbox"].setChecked(True)
                # Set severity based on count
                if risk_key == "awol":
                    # AWOL: any occurrence = HIGH
                    self._current_widgets[risk_key]["slider"].setValue(3)  # HIGH
                elif count >= 3:
                    self._current_widgets[risk_key]["slider"].setValue(3)  # HIGH
                else:
                    self._current_widgets[risk_key]["slider"].setValue(2)  # MEDIUM

        # Auto-populate Historical Risk checkboxes (ALL time)
        # Severity based on average per year: > 5/year = HIGH, <= 5/year = MEDIUM, < 1/year = LOW
        # If risk is present currently, it's also present historically
        for risk_key, count in historical_risk_counts.items():
            if risk_key in self._historical_widgets and count > 0:
                self._historical_widgets[risk_key]["checkbox"].setChecked(True)
                # Calculate average per year
                avg_per_year = count / num_years
                # Set severity based on average per year
                if avg_per_year > 5:
                    self._historical_widgets[risk_key]["slider"].setValue(3)  # HIGH
                elif avg_per_year < 1:
                    self._historical_widgets[risk_key]["slider"].setValue(1)  # LOW
                else:
                    self._historical_widgets[risk_key]["slider"].setValue(2)  # MEDIUM

        if not all_incidents:
            self.extracted_section.setVisible(False)
            return

        # Sort by date (newest first)
        sorted_incidents = sorted(all_incidents, key=lambda x: x["date"] or datetime.min, reverse=True)

        # Store all incidents for filtering
        self._all_risk_incidents = sorted_incidents
        self._current_risk_filter = None

        # Risk type colors for visual indication
        self._risk_colors = {
            "Verbal Aggression": "#9E9E9E",
            "Physical Aggression": "#b71c1c",
            "Self-Harm": "#e91e63",
            "Suicide": "#9c27b0",
            "Absconding": "#2196f3",
            "AWOL/Absconding": "#2196f3",
            "Property Damage": "#ff9800",
            "Sexual Behaviour": "#673ab7",
            "Vulnerability": "#00bcd4",
            "Self-Neglect": "#795548",
            "Non-Compliance": "#607d8b",
            "Bullying/Exploitation": "#009688",
            "Substance Misuse": "#9c27b0",
        }

        self._severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        # Collect unique labels (category or category:subcategory)
        labels = {}
        for inc in sorted_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat:
                label = f"{cat}: {subcat}"
            else:
                label = cat
            if label not in labels:
                labels[label] = self._risk_colors.get(cat, "#666666")

        # Create filter panel with horizontal scroll
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(40)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
            QScrollBar:horizontal {
                height: 6px;
                background: #f0f0f0;
                border-radius: 3px;
            }
            QScrollBar::handle:horizontal {
                background: #c0c0c0;
                border-radius: 3px;
                min-width: 20px;
            }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-size: 16px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background: {color}dd;
                }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text: self._apply_risk_filter(lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row (hidden initially)
        self._risk_filter_status_widget = QWidget()
        self._risk_filter_status_widget.setStyleSheet("background: transparent;")
        self._risk_filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(self._risk_filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        self._risk_filter_label = QLabel("Filtered by: ")
        self._risk_filter_label.setStyleSheet("font-size: 17px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(self._risk_filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 4px 10px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #dc2626;
            }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(self._remove_risk_filter)
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(self._risk_filter_status_widget)
        self.extracted_checkboxes_layout.addWidget(filter_container)

        # Container for incident entries (for re-rendering on filter)
        self._risk_incidents_container = QWidget()
        self._risk_incidents_container.setStyleSheet("background: transparent;")
        self._risk_incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        self._risk_incidents_layout = QVBoxLayout(self._risk_incidents_container)
        self._risk_incidents_layout.setContentsMargins(0, 0, 0, 0)
        self._risk_incidents_layout.setSpacing(8)
        self.extracted_checkboxes_layout.addWidget(self._risk_incidents_container)

        # Render all incidents
        self._render_risk_incidents(sorted_incidents)

        # Show section
        self.extracted_section.setVisible(True)
        # Keep collapsed on open
        # if self.extracted_section._is_collapsed:
        #     self.extracted_section._toggle_collapse()

    def _apply_risk_filter(self, label: str):
        """Apply filter to show only incidents matching the label."""
        self._current_risk_filter = label
        self._risk_filter_label.setText(f"Filtered by: {label}")
        self._risk_filter_status_widget.setVisible(True)

        # Filter incidents
        filtered = []
        for inc in self._all_risk_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat:
                inc_label = f"{cat}: {subcat}"
            else:
                inc_label = cat
            if inc_label == label:
                filtered.append(inc)

        self._render_risk_incidents(filtered)

    def _remove_risk_filter(self):
        """Remove filter and show all incidents."""
        self._current_risk_filter = None
        self._risk_filter_status_widget.setVisible(False)
        self._render_risk_incidents(self._all_risk_incidents)

    def _render_risk_incidents(self, incidents: list):
        """Render the list of risk incidents."""
        import re
        import html
        from PySide6.QtWidgets import QTextEdit
        from datetime import datetime

        # Clear existing
        while self._risk_incidents_layout.count():
            child = self._risk_incidents_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Clear checkboxes list (only the ones in incidents container)
        self._extracted_checkboxes.clear()

        for incident in incidents:
            date = incident["date"]
            cat_name = incident["category"]
            text = incident["text"]
            matched = incident["matched"]
            subcat_name = incident["subcategory"]
            severity = incident["severity"]

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Get colors
            cat_color = self._risk_colors.get(cat_name, "#666666")
            sev_color = self._severity_colors.get(severity, "#666666")

            # Create HTML with highlighted matched text
            escaped_text = html.escape(text)
            if matched:
                # Escape the matched text and create a case-insensitive pattern
                escaped_matched = html.escape(matched)
                try:
                    # Highlight all occurrences of the matched text (case-insensitive)
                    pattern = re.compile(re.escape(escaped_matched), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_matched}</span>',
                        escaped_text
                    )
                except:
                    highlighted_html = escaped_text
            else:
                highlighted_html = escaped_text

            # Wrap in HTML body with styling
            full_html = f'''
            <html>
            <body style="font-family: -apple-system, BlinkMacSystemFont, sans-serif; font-size: 17px; color: #333; margin: 0; padding: 0;">
            {highlighted_html}
            </body>
            </html>
            '''

            # Create collapsible entry box with risk type indicator
            from PySide6.QtWidgets import QTextEdit
            entry_frame = QFrame()
            entry_frame.setObjectName("riskEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#riskEntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-left: 4px solid {cat_color};
                    border-radius: 8px;
                    padding: 2px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 6, 6, 6)
            entry_layout.setSpacing(4)

            # Header row: toggle → date → risk badge → severity → stretch → checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button on the LEFT
            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(180, 150, 50, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: #806000;
                }
                QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 500;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Risk type badge
            risk_badge = QLabel(f"{cat_name}: {subcat_name}")
            risk_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 14px;
                    font-weight: 600;
                    color: white;
                    background: {cat_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 6px;
                }}
            """)
            header_row.addWidget(risk_badge)

            # Severity badge
            sev_badge = QLabel(severity.upper())
            sev_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 13px;
                    font-weight: 700;
                    color: white;
                    background: {sev_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 5px;
                }}
            """)
            header_row.addWidget(sev_badge)
            header_row.addStretch()

            # Checkbox on the RIGHT
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(self._update_preview)
            header_row.addWidget(cb)
            self._extracted_checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body (full text with highlighted match, hidden by default)
            body_text = QTextEdit()
            body_text.setHtml(full_html)
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            body_text.setStyleSheet("""
                QTextEdit {
                    font-size: 16px;
                    color: #333;
                    background: rgba(255, 248, 220, 0.5);
                    border: none;
                    padding: 6px;
                    border-radius: 4px;
                }
            """)
            # Set minimum and maximum height
            body_text.setMinimumHeight(50)
            body_text.setMaximumHeight(150)
            body_text.setVisible(False)
            entry_layout.addWidget(body_text)

            drag_bar = QFrame()
            drag_bar.setFixedHeight(8)
            drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
            drag_bar.setStyleSheet("""
                QFrame {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                    border-radius: 2px; margin: 2px 40px;
                }
                QFrame:hover {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                }
            """)
            drag_bar.setVisible(False)
            drag_bar._drag_y = None
            drag_bar._init_h = None
            def _make_drag_handlers(handle, text_widget):
                def press(ev):
                    handle._drag_y = ev.globalPosition().y()
                    handle._init_h = text_widget.height()
                def move(ev):
                    if handle._drag_y is not None:
                        delta = int(ev.globalPosition().y() - handle._drag_y)
                        new_h = max(50, handle._init_h + delta)
                        text_widget.setMinimumHeight(new_h)
                        text_widget.setMaximumHeight(new_h)
                def release(ev):
                    if handle._drag_y is not None:
                        text_widget.setMaximumHeight(16777215)
                        handle._drag_y = None
                return press, move, release
            dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
            drag_bar.mousePressEvent = dp
            drag_bar.mouseMoveEvent = dm
            drag_bar.mouseReleaseEvent = dr
            entry_layout.addWidget(drag_bar)

            # Toggle function
            def make_toggle(btn, body, frame, popup_self, bar):
                def toggle():
                    is_visible = body.isVisible()
                    body.setVisible(not is_visible)
                    bar.setVisible(not is_visible)
                    btn.setText("▾" if not is_visible else "▸")
                    frame.updateGeometry()
                    if hasattr(popup_self, 'extracted_container'):
                        popup_self.extracted_container.updateGeometry()
                        popup_self.extracted_container.update()
                return toggle

            toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
            toggle_btn.clicked.connect(toggle_fn)
            date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

            self._risk_incidents_layout.addWidget(entry_frame)

        self._risk_incidents_layout.addStretch()


# ================================================================
# GPR BACKGROUND POPUP (wrapper around BackgroundHistoryPopup)
# ================================================================

class GPRBackgroundPopup(QWidget):
    """Wrapper around BackgroundHistoryPopup that adds risk in childhood dropdown and adapts for GPR."""

    sent = Signal(str)

    def __init__(self, parent=None, gender: str = None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._entries = []
        self._risk_entries = []
        self._risk_childhood_text = ""
        self._gender = gender
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Use the existing BackgroundHistoryPopup with gender
        self.background_popup = BackgroundHistoryPopup(gender=self._gender, parent=self)

        # Prevent BackgroundHistoryPopup from closing itself after send
        # (it normally calls self.close() which hides the widget and breaks the GPR layout)
        self.background_popup.close = lambda: None

        # Connect the signal (BackgroundHistoryPopup emits (str, dict), we just want str)
        self.background_popup.sent.connect(self._on_background_sent)

        # Inject "Risk history in childhood" dropdown at the top of Family & Childhood section
        self._inject_risk_childhood_dropdown()

        layout.addWidget(self.background_popup)

    def _inject_risk_childhood_dropdown(self):
        """Inject the Risk history in childhood dropdown into the Family & Childhood section."""
        # Find the family_history_widget in the BackgroundHistoryPopup
        # and insert our dropdown before it
        family_widget = self.background_popup.family_history_widget

        # Find the parent layout (the form_layout inside the scroll container)
        # Navigate up: family_history_widget -> column layout -> form_layout
        parent_widget = family_widget.parentWidget()
        if parent_widget:
            form_layout = parent_widget.layout()
            if form_layout:
                # Find the index of the family history column layout
                for i in range(form_layout.count()):
                    item = form_layout.itemAt(i)
                    if item and item.layout():
                        # Check if this layout contains our family_history_widget
                        inner_layout = item.layout()
                        for j in range(inner_layout.count()):
                            inner_item = inner_layout.itemAt(j)
                            if inner_item and inner_item.widget() == family_widget:
                                # Found it! Insert our dropdown before this index
                                self._create_risk_dropdown(form_layout, i)
                                return

    def _create_risk_dropdown(self, form_layout, insert_index):
        """Create the Risk history in childhood dropdown and insert it into the form layout."""
        # Create a column layout similar to BackgroundHistoryPopup's column() function
        col = QVBoxLayout()
        col.setSpacing(4)
        col.setAlignment(Qt.AlignmentFlag.AlignTop)

        # Label
        lbl = QLabel("Risk history in childhood")
        lbl.setStyleSheet("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #003c32;
                margin: 0px;
                padding: 0px 0px 4px 0px;
                background: transparent;
                border: none;
            }
        """)
        lbl.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
        col.addWidget(lbl)

        # Dropdown
        self.risk_childhood_combo = NoWheelComboBox()
        self.risk_childhood_combo.setEditable(True)
        self.risk_childhood_combo.setStyleSheet("""
            QComboBox {
                background: white;
                color: #333333;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 6px 10px;
                min-height: 28px;
            }
            QComboBox::drop-down {
                border: none;
                width: 24px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #dbeafe;
            }
        """)

        # Default options
        self.risk_childhood_combo.addItems([
            "",
            "No significant risk behavior in childhood reported",
            "History of conduct disorder in childhood",
            "History of aggressive behavior in childhood",
            "History of fire-setting in childhood",
            "History of cruelty to animals in childhood",
            "History of truancy and school refusal",
            "History of substance misuse from early age",
            "Early involvement with criminal justice system",
        ])

        # Connect to update internal state
        self.risk_childhood_combo.currentTextChanged.connect(self._on_risk_childhood_changed)
        col.addWidget(self.risk_childhood_combo)

        # Insert into the form layout before the family history section
        form_layout.insertLayout(insert_index, col)

    def _on_risk_childhood_changed(self, text: str):
        """Store the risk in childhood selection and trigger send."""
        self._risk_childhood_text = text
        # Trigger a re-send with the updated risk text
        self.background_popup._refresh_preview()

    def _on_background_sent(self, text: str, state: dict):
        """Adapt the signal from BackgroundHistoryPopup to just emit the text, appending risk in childhood if set."""
        # Append risk in childhood text at the end
        if self._risk_childhood_text and self._risk_childhood_text.strip():
            text = f"{text}\n\n{self._risk_childhood_text}" if text else self._risk_childhood_text
        self.sent.emit(text)

    def set_entries(self, entries: list):
        """Set entries from extracted data, preserving dates."""
        self._entries = entries

        # Pass entries directly to set_extracted_data which supports list of dicts with 'date' and 'text'
        if entries:
            self.background_popup.set_extracted_data(entries)

    def set_risk_entries(self, entries: list):
        """Set risk-specific entries that can populate the risk childhood dropdown."""
        self._risk_entries = entries

        # If we have risk entries, add them to the dropdown
        if entries and hasattr(self, 'risk_childhood_combo'):
            for entry in entries:
                text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
                if text and text.strip():
                    # Add to combo if not already present
                    if self.risk_childhood_combo.findText(text) < 0:
                        self.risk_childhood_combo.addItem(text[:200])  # Truncate if too long

    def update_gender(self, gender: str):
        """Update gender for pronoun generation in the background popup."""
        self._gender = gender
        if hasattr(self, 'background_popup') and hasattr(self.background_popup, 'update_gender'):
            self.background_popup.update_gender(gender)


# ================================================================
# GPR MEDICAL HISTORY POPUP (wrapper around PhysicalHealthPopup)
# ================================================================

class GPRMedicalHistoryPopup(QWidget):
    """Wrapper around PhysicalHealthPopup with collapsible sections and drag bars."""

    sent = Signal(str)

    def __init__(self, parent=None, gender: str = None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._entries = []
        self._gender = gender or "Male"
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: INPUT FORM (collapsible with drag bar)
        # ====================================================
        self.input_section = CollapsibleSection("Physical Health Conditions", start_collapsed=True)
        self.input_section.set_content_height(300)
        self.input_section._min_height = 100
        self.input_section._max_height = 400
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)

        # Use the existing PhysicalHealthPopup
        self.health_popup = PhysicalHealthPopup(gender=self._gender, parent=self)

        # Prevent PhysicalHealthPopup from closing itself after send
        self.health_popup.close = lambda: None

        # Connect to the health popup's sent signal to forward all changes
        self.health_popup.sent.connect(self._on_health_sent)

        # Also connect checkbox changes to update our main preview (for extracted data too)
        for category, cb_list in self.health_popup._checkboxes.items():
            for cb in cb_list:
                cb.stateChanged.connect(self._update_main_preview)

        self.input_section.set_content(self.health_popup)
        main_layout.addWidget(self.input_section)

        # ====================================================
        # SECTION 3: IMPORTED DATA (matching BackgroundHistoryPopup)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(150)
        self.extracted_section._min_height = 60
        self.extracted_section._max_height = 400
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        extracted_content = QWidget()
        extracted_content.setStyleSheet("""
            QWidget {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
            QCheckBox {
                background: transparent;
                border: none;
                padding: 4px;
                font-size: 17px;
                color: #4a4a4a;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
            }
        """)

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollArea > QWidget > QWidget { background: transparent; }
        """)

        # Container for checkboxes
        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)
        self.extracted_checkboxes_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        # Store extracted checkboxes
        self._extracted_checkboxes = []

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _update_main_preview(self):
        """Build text from selections and send directly to card."""
        input_parts = []
        extracted_parts = []

        # Get selected conditions from the health popup checkboxes
        selected_conditions = []
        for category, cb_list in self.health_popup._checkboxes.items():
            for cb in cb_list:
                if cb.isChecked():
                    selected_conditions.append(cb.text())

        if selected_conditions:
            # Generate formatted text using the health popup's method
            text = self.health_popup.formatted_text(selected_conditions)
            input_parts.append(text)

        # Add checked extracted entries
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                extracted_parts.append(cb.property("full_text"))

        # Combine with blank line separator between sections
        if input_parts and extracted_parts:
            combined = "\n\n".join(input_parts) + "\n\n" + "\n\n".join(extracted_parts)
        elif input_parts:
            combined = "\n\n".join(input_parts)
        elif extracted_parts:
            combined = "\n\n".join(extracted_parts)
        else:
            combined = ""

        # Send directly to card
        self.sent.emit(combined)

    def _on_health_sent(self, text: str, state: dict):
        """Handle signal from PhysicalHealthPopup - use our combined preview to include extracted data."""
        # Don't emit directly - let _update_main_preview handle it to include extracted data
        # self.sent.emit(text)
        pass  # _update_main_preview is already connected to checkbox stateChanged

    def update_gender(self, gender: str):
        """Update gender for pronoun generation in the health popup."""
        self._gender = gender
        if hasattr(self, 'health_popup') and hasattr(self.health_popup, 'update_gender'):
            self.health_popup.update_gender(gender)
        # Also refresh our own preview to regenerate text with new pronouns
        self._update_main_preview()

    def set_entries(self, items: list):
        """Display extracted data with collapsible dated entry boxes (matching BackgroundHistoryPopup)."""
        self._entries = items

        # Clear existing checkboxes
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        # Clear layout
        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Handle legacy string format
        if isinstance(items, str):
            if items.strip():
                items = [{"date": None, "text": p.strip()} for p in items.split("\n\n") if p.strip()]
            else:
                items = []

        if items:
            # Sort by date (newest first)
            def get_sort_date(item):
                dt = item.get("date")
                if dt is None:
                    return ""
                if hasattr(dt, "strftime"):
                    return dt.strftime("%Y-%m-%d")
                return str(dt)

            sorted_items = sorted(items, key=get_sort_date, reverse=True)

            for item in sorted_items:
                dt = item.get("date")
                text = item.get("text", "").strip()
                if not text:
                    continue

                # Format date for header
                if dt:
                    if hasattr(dt, "strftime"):
                        date_str = dt.strftime("%d %b %Y")
                    else:
                        date_str = str(dt)
                else:
                    date_str = "No date"

                # Create collapsible entry box
                entry_frame = QFrame()
                entry_frame.setObjectName("entryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
                entry_frame.setStyleSheet("""
                    QFrame#entryFrame {
                        background: rgba(255, 255, 255, 0.95);
                        border: 1px solid rgba(180, 150, 50, 0.4);
                        border-radius: 8px;
                        padding: 4px;
                    }
                """)
                entry_layout = QVBoxLayout(entry_frame)
                entry_layout.setContentsMargins(10, 8, 10, 8)
                entry_layout.setSpacing(6)
                entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

                # Header row with toggle, date label, and checkbox
                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                # Toggle button on the LEFT
                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(180, 150, 50, 0.2);
                        border: none;
                        border-radius: 4px;
                        font-size: 17px;
                        font-weight: bold;
                        color: #806000;
                    }
                    QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
                """)
                header_row.addWidget(toggle_btn)

                # Date label
                from PySide6.QtWidgets import QTextEdit
                date_label = QLabel(f"📅 {date_str}")
                date_label.setStyleSheet("""
                    QLabel {
                        font-size: 17px;
                        font-weight: 600;
                        color: #806000;
                        background: transparent;
                        border: none;
                    }
                """)
                date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                header_row.addWidget(date_label)
                header_row.addStretch()

                # Checkbox on the RIGHT
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("""
                    QCheckBox { background: transparent; }
                    QCheckBox::indicator { width: 16px; height: 16px; }
                """)
                cb.stateChanged.connect(self._update_main_preview)
                header_row.addWidget(cb)

                entry_layout.addLayout(header_row)

                # Body (full content, hidden by default)
                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setStyleSheet("""
                    QTextEdit {
                        font-size: 17px;
                        color: #333;
                        background: rgba(255, 248, 220, 0.5);
                        border: none;
                        padding: 8px;
                        border-radius: 6px;
                    }
                """)
                # Calculate height based on content
                body_text.document().setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
                doc_height = body_text.document().size().height() + 20
                body_text.setFixedHeight(int(max(doc_height, 60)))
                body_text.setVisible(False)
                entry_layout.addWidget(body_text)

                # Toggle function
                def make_toggle(btn, body, frame, popup_self):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                        if hasattr(popup_self, 'extracted_container'):
                            popup_self.extracted_container.updateGeometry()
                            popup_self.extracted_container.update()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self)
                toggle_btn.clicked.connect(toggle_fn)
                date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                self.extracted_checkboxes_layout.addWidget(entry_frame)
                self._extracted_checkboxes.append(cb)

            self.extracted_section.setVisible(True)
            # Keep collapsed on open
            # if self.extracted_section._is_collapsed:
            #     self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)


# ================================================================
# GPR SUBSTANCE USE POPUP (wrapper around DrugsAlcoholPopup)
# ================================================================

class GPRSubstanceUsePopup(QWidget):
    """Wrapper around DrugsAlcoholPopup with collapsible sections and extracted data."""

    sent = Signal(str)

    def __init__(self, parent=None, gender: str = None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._entries = []
        self._extracted_checkboxes = []
        self._gender = gender or "Male"
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: INPUT FORM (collapsible with drag bar)
        # ====================================================
        self.input_section = CollapsibleSection("Substance Use", start_collapsed=True)
        self.input_section.set_content_height(350)
        self.input_section._min_height = 100
        self.input_section._max_height = 500
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.input_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #005a50;
                background: transparent;
                border: none;
            }
        """)

        # Use the existing DrugsAlcoholPopup
        self.substance_popup = DrugsAlcoholPopup(gender=self._gender, parent=self)
        self.substance_popup.setStyleSheet("""
            QWidget { background: rgba(255,255,255,0.95); }
            QLabel { background: transparent; color: #2b2b2b; }
        """)

        # Prevent DrugsAlcoholPopup from closing itself after send
        self.substance_popup.close = lambda: None

        # Connect slider changes to update our main preview
        for slider in [self.substance_popup.alc_age, self.substance_popup.alc_amt,
                       self.substance_popup.smoke_age, self.substance_popup.smoke_amt,
                       self.substance_popup.drug_age, self.substance_popup.drug_amt]:
            slider.slider.valueChanged.connect(self._update_main_preview)

        # Connect drug radio buttons
        for rb in self.substance_popup.drug_buttons.values():
            rb.toggled.connect(self._update_main_preview)

        self.input_section.set_content(self.substance_popup)
        main_layout.addWidget(self.input_section)

        # ====================================================
        # SECTION 3: IMPORTED DATA (collapsible, at bottom)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(300)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 500
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        # Use QScrollArea directly as content
        self.extracted_scroll = QScrollArea()
        self.extracted_scroll.setWidgetResizable(True)
        self.extracted_scroll.setFrameShape(QScrollArea.NoFrame)
        self.extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.extracted_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        # Container for checkboxes inside scroll area
        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(12, 12, 12, 12)
        self.extracted_checkboxes_layout.setSpacing(12)

        self.extracted_scroll.setWidget(self.extracted_container)
        self.extracted_section.set_content(self.extracted_scroll)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _update_main_preview(self):
        """Build text from selections and send directly to card."""
        input_parts = []
        extracted_parts = []

        # Get formatted text from the substance popup
        text = self.substance_popup.formatted_text()
        if text:
            input_parts.append(text)

        # Get checked extracted entries
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                extracted_parts.append(cb.property("full_text"))

        # Combine with blank line separator
        if input_parts and extracted_parts:
            combined = "\n\n".join(input_parts) + "\n\n" + "\n\n".join(extracted_parts)
        elif input_parts:
            combined = "\n\n".join(input_parts)
        elif extracted_parts:
            combined = "\n\n".join(extracted_parts)
        else:
            combined = ""

        # Send directly to card
        self.sent.emit(combined)

    def update_gender(self, gender: str):
        """Update gender for pronoun generation in the substance popup."""
        self._gender = gender
        if hasattr(self, 'substance_popup') and hasattr(self.substance_popup, 'update_gender'):
            self.substance_popup.update_gender(gender)
        # Also refresh our own preview to regenerate text with new pronouns
        self._update_main_preview()

    def set_patient_age(self, age: int):
        """Set patient age to limit age sliders appropriately."""
        if hasattr(self, 'substance_popup') and hasattr(self.substance_popup, 'set_patient_age'):
            self.substance_popup.set_patient_age(age)

    def set_entries(self, entries: list):
        """Set entries from extracted data with checkboxes."""
        self._entries = entries

        # Clear existing checkboxes
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        # Clear layout
        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        has_entries = False
        for entry in entries:
            text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
            if text and text.strip():
                has_entries = True

                # Create a container for checkbox + label
                entry_widget = QWidget()
                entry_widget.setStyleSheet("background: transparent;")
                entry_layout = QHBoxLayout(entry_widget)
                entry_layout.setContentsMargins(0, 4, 0, 4)
                entry_layout.setSpacing(8)
                entry_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

                # Checkbox
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(20, 20)
                cb.setStyleSheet("""
                    QCheckBox {
                        background: transparent;
                    }
                    QCheckBox::indicator {
                        width: 16px;
                        height: 16px;
                    }
                """)
                cb.stateChanged.connect(self._update_main_preview)
                entry_layout.addWidget(cb, 0, Qt.AlignmentFlag.AlignTop)

                # Label with word wrap
                label = QLabel(text)
                label.setWordWrap(True)
                label.setStyleSheet("""
                    QLabel {
                        background: transparent;
                        border: none;
                        font-size: 17px;
                        color: #4a4a4a;
                        padding: 0px;
                    }
                """)
                entry_layout.addWidget(label, 1)

                self.extracted_checkboxes_layout.addWidget(entry_widget)
                self._extracted_checkboxes.append(cb)

        if has_entries:
            self.extracted_section.setVisible(True)
            if self.extracted_section.is_collapsed():
                self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)

    def set_notes_for_substance_analysis(self, notes: list, extracted_entries: list = None):
        """Run risk analysis on notes to find substance misuse and combine with data extractor findings.

        Display in date order with the same format as section 7 (Risk).
        """
        from risk_overview_panel import analyze_notes_for_risk
        from datetime import datetime
        import re
        import html

        # Clear existing checkboxes
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        # Clear layout
        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        all_incidents = []

        # Run risk analysis on notes to get substance misuse incidents
        if notes:
            results = analyze_notes_for_risk(notes)

            # Only get "Substance Misuse" category
            substance_data = results.get("categories", {}).get("Substance Misuse", {})
            for incident in substance_data.get("incidents", []):
                all_incidents.append({
                    "date": incident.get("date"),
                    "text": incident.get("full_text", ""),
                    "matched": incident.get("matched", ""),
                    "subcategory": incident.get("subcategory", ""),
                    "severity": incident.get("severity", "medium"),
                    "source": "notes",
                })

        # Add data extractor entries - integrate with risk analysis in date order
        if extracted_entries:
            print(f"[SUBSTANCE] Processing {len(extracted_entries)} extracted entries")
            for i, entry in enumerate(extracted_entries):
                if isinstance(entry, dict):
                    text = entry.get("text", "")
                    date = entry.get("date")
                    print(f"[SUBSTANCE] Entry {i}: date={date}, text={text[:50] if text else 'None'}...")
                else:
                    text = str(entry)
                    date = None
                    print(f"[SUBSTANCE] Entry {i} is string: {text[:50]}...")
                if text and text.strip():
                    all_incidents.append({
                        "date": date,
                        "text": text,
                        "matched": "",
                        "subcategory": "Data Extractor",
                        "severity": "medium",
                        "source": "extractor",
                    })
                    print(f"[SUBSTANCE] Added entry {i} with date={date}")

        if not all_incidents:
            self.extracted_section.setVisible(False)
            return

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                return datetime.min  # Undated at end
            return d

        sorted_incidents = sorted(all_incidents, key=get_sort_date, reverse=True)

        # Store all incidents for filtering
        self._all_substance_incidents = sorted_incidents
        self._current_substance_filter = None

        # Count by source
        from_notes = sum(1 for x in sorted_incidents if x.get("source") == "notes")
        from_extractor = sum(1 for x in sorted_incidents if x.get("source") == "extractor")
        print(f"[SUBSTANCE] ===== FINAL: {len(sorted_incidents)} total incidents =====")
        print(f"[SUBSTANCE]   From risk analysis: {from_notes}")
        print(f"[SUBSTANCE]   From data extractor: {from_extractor}")

        # Subcategory colors
        self._substance_colors = {
            "Positive Drug Test": "#6a1b9a",
            "Smelling of Substances": "#7b1fa2",
            "Appeared Intoxicated": "#8e24aa",
            "Admitted Substance Use": "#9c27b0",
            "Found with Substances": "#ab47bc",
            "Data Extractor": "#607d8b",
        }

        self._substance_severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        # Collect unique labels (subcategory)
        labels = {}
        for inc in sorted_incidents:
            subcat = inc.get("subcategory", "")
            label = f"Substance Misuse: {subcat}" if subcat else "Substance Misuse"
            if label not in labels:
                labels[label] = self._substance_colors.get(subcat, "#9c27b0")

        # Create filter panel with horizontal scroll
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(40)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
            QScrollBar:horizontal {
                height: 6px;
                background: #f0f0f0;
                border-radius: 3px;
            }
            QScrollBar::handle:horizontal {
                background: #c0c0c0;
                border-radius: 3px;
                min-width: 20px;
            }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-size: 16px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background: {color}dd;
                }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text: self._apply_substance_filter(lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row (hidden initially)
        self._substance_filter_status_widget = QWidget()
        self._substance_filter_status_widget.setStyleSheet("background: transparent;")
        self._substance_filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(self._substance_filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        self._substance_filter_label = QLabel("Filtered by: ")
        self._substance_filter_label.setStyleSheet("font-size: 17px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(self._substance_filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 4px 10px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #dc2626;
            }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(self._remove_substance_filter)
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(self._substance_filter_status_widget)
        self.extracted_checkboxes_layout.addWidget(filter_container)

        # Container for incident entries (for re-rendering on filter)
        self._substance_incidents_container = QWidget()
        self._substance_incidents_container.setStyleSheet("background: transparent;")
        self._substance_incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        self._substance_incidents_layout = QVBoxLayout(self._substance_incidents_container)
        self._substance_incidents_layout.setContentsMargins(0, 0, 0, 0)
        self._substance_incidents_layout.setSpacing(8)
        self.extracted_checkboxes_layout.addWidget(self._substance_incidents_container)

        # Render all incidents
        self._render_substance_incidents(sorted_incidents)

        # Show section
        self.extracted_section.setVisible(True)
        if self.extracted_section.is_collapsed():
            self.extracted_section._toggle_collapse()

    def _apply_substance_filter(self, label: str):
        """Apply filter to show only incidents matching the label."""
        self._current_substance_filter = label
        self._substance_filter_label.setText(f"Filtered by: {label}")
        self._substance_filter_status_widget.setVisible(True)

        # Filter incidents
        filtered = []
        for inc in self._all_substance_incidents:
            subcat = inc.get("subcategory", "")
            inc_label = f"Substance Misuse: {subcat}" if subcat else "Substance Misuse"
            if inc_label == label:
                filtered.append(inc)

        self._render_substance_incidents(filtered)

    def _remove_substance_filter(self):
        """Remove filter and show all incidents."""
        self._current_substance_filter = None
        self._substance_filter_status_widget.setVisible(False)
        self._render_substance_incidents(self._all_substance_incidents)

    def _render_substance_incidents(self, incidents: list):
        """Render the list of substance incidents."""
        import re
        import html
        from PySide6.QtWidgets import QTextEdit, QFrame
        from datetime import datetime

        # Clear existing
        while self._substance_incidents_layout.count():
            child = self._substance_incidents_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Clear checkboxes list
        self._extracted_checkboxes.clear()

        for incident in incidents:
            date = incident["date"]
            text = incident["text"]
            matched = incident["matched"]
            subcat_name = incident["subcategory"]
            severity = incident["severity"]
            source = incident["source"]

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Get colors
            subcat_color = self._substance_colors.get(subcat_name, "#9c27b0")
            sev_color = self._substance_severity_colors.get(severity, "#666666")

            # Create HTML with highlighted matched text
            escaped_text = html.escape(text)
            if matched:
                escaped_matched = html.escape(matched)
                try:
                    pattern = re.compile(re.escape(escaped_matched), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_matched}</span>',
                        escaped_text
                    )
                except:
                    highlighted_html = escaped_text
            else:
                highlighted_html = escaped_text

            full_html = f'''
            <html>
            <body style="font-family: -apple-system, BlinkMacSystemFont, sans-serif; font-size: 17px; color: #333; margin: 0; padding: 0;">
            {highlighted_html}
            </body>
            </html>
            '''

            # Create entry frame with colored left border
            entry_frame = QFrame()
            entry_frame.setObjectName("substanceEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#substanceEntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-left: 4px solid {subcat_color};
                    border-radius: 8px;
                    padding: 2px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 6, 6, 6)
            entry_layout.setSpacing(4)

            # Header row: toggle → date → subcategory badge → severity → stretch → checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button
            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(180, 150, 50, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: #806000;
                }
                QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 500;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)
            header_row.addWidget(date_label)

            # Subcategory badge
            badge_text = f"Substance Misuse: {subcat_name}" if subcat_name else "Substance Misuse"
            subcat_badge = QLabel(badge_text)
            subcat_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 14px;
                    font-weight: 600;
                    color: white;
                    background: {subcat_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 6px;
                }}
            """)
            header_row.addWidget(subcat_badge)

            # Severity badge
            sev_badge = QLabel(severity.upper())
            sev_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 13px;
                    font-weight: 700;
                    color: white;
                    background: {sev_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 5px;
                }}
            """)
            header_row.addWidget(sev_badge)
            header_row.addStretch()

            # Checkbox on the RIGHT
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(self._update_main_preview)
            header_row.addWidget(cb)
            self._extracted_checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Text content
            body_text = QTextEdit()
            body_text.setReadOnly(True)
            body_text.setHtml(full_html)
            body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            body_text.setMinimumHeight(60)
            body_text.setMaximumHeight(120)
            body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            body_text.setStyleSheet("""
                QTextEdit {
                    background: rgba(255, 248, 220, 0.5);
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                    font-size: 17px;
                    color: #4a4a4a;
                }
            """)
            body_text.setVisible(False)
            entry_layout.addWidget(body_text)

            drag_bar = QFrame()
            drag_bar.setFixedHeight(8)
            drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
            drag_bar.setStyleSheet("""
                QFrame {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                    border-radius: 2px; margin: 2px 40px;
                }
                QFrame:hover {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                }
            """)
            drag_bar.setVisible(False)
            drag_bar._drag_y = None
            drag_bar._init_h = None
            def _make_drag_handlers(handle, text_widget):
                def press(ev):
                    handle._drag_y = ev.globalPosition().y()
                    handle._init_h = text_widget.height()
                def move(ev):
                    if handle._drag_y is not None:
                        delta = int(ev.globalPosition().y() - handle._drag_y)
                        new_h = max(60, handle._init_h + delta)
                        text_widget.setMinimumHeight(new_h)
                        text_widget.setMaximumHeight(new_h)
                def release(ev):
                    if handle._drag_y is not None:
                        text_widget.setMaximumHeight(16777215)
                        handle._drag_y = None
                return press, move, release
            dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
            drag_bar.mousePressEvent = dp
            drag_bar.mouseMoveEvent = dm
            drag_bar.mouseReleaseEvent = dr
            entry_layout.addWidget(drag_bar)

            def make_toggle(btn, body, bar):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        bar.setVisible(False)
                        btn.setText("▸")
                    else:
                        body.setVisible(True)
                        bar.setVisible(True)
                        btn.setText("▾")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text, drag_bar))
            date_label.mousePressEvent = lambda e, btn=toggle_btn: btn.click()

            self._substance_incidents_layout.addWidget(entry_frame)

        self._substance_incidents_layout.addStretch()


# ================================================================
# GPR MEDICATION POPUP
# ================================================================

FREQUENCY_OPTIONS = ["OD", "BD", "TDS", "QDS", "Nocte", "PRN", "Weekly", "Fortnightly", "Monthly"]

class GPRMedicationPopup(QWidget):
    """Medication popup with preview, input section, and extracted data sections."""

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._medications = []
        self._extracted_checkboxes = []
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        from PySide6.QtWidgets import QComboBox, QFrame

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: MEDICATION INPUT (collapsible with drag bar)
        # ====================================================
        self.input_section = CollapsibleSection("Current Medication", start_collapsed=True)
        self.input_section.set_content_height(350)
        self.input_section._min_height = 100
        self.input_section._max_height = 500
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.input_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #005a50;
                background: transparent;
                border: none;
            }
        """)

        # Medication form container
        med_form_scroll = QScrollArea()
        med_form_scroll.setWidgetResizable(True)
        med_form_scroll.setFrameShape(QScrollArea.NoFrame)
        med_form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        med_form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        med_form_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255,255,255,0.95);
                border: 1px solid rgba(0, 140, 126, 0.2);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        med_form_container = QWidget()
        med_form_container.setStyleSheet("background: transparent;")
        med_form_layout = QVBoxLayout(med_form_container)
        med_form_layout.setContentsMargins(12, 12, 12, 12)
        med_form_layout.setSpacing(8)

        # Medication entries container
        self.med_entries_container = QWidget()
        self.med_entries_container.setStyleSheet("background: transparent;")
        self.med_entries_layout = QVBoxLayout(self.med_entries_container)
        self.med_entries_layout.setContentsMargins(0, 0, 0, 0)
        self.med_entries_layout.setSpacing(8)
        med_form_layout.addWidget(self.med_entries_container)

        # Add first medication entry
        self._add_medication_entry()

        # Add medication button
        add_med_btn = QPushButton("+ Add Medication")
        add_med_btn.setStyleSheet("""
            QPushButton {
                background: #e5e7eb;
                color: #374151;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #d1d5db;
            }
        """)
        add_med_btn.clicked.connect(self._add_medication_entry)
        med_form_layout.addWidget(add_med_btn)

        med_form_layout.addStretch()
        med_form_scroll.setWidget(med_form_container)
        self.input_section.set_content(med_form_scroll)
        main_layout.addWidget(self.input_section)

        # ====================================================
        # SECTION 3: IMPORTED DATA (collapsible, at bottom)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 400
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        # Use QScrollArea directly as content
        self.extracted_scroll = QScrollArea()
        self.extracted_scroll.setWidgetResizable(True)
        self.extracted_scroll.setFrameShape(QScrollArea.NoFrame)
        self.extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.extracted_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        # Container for checkboxes inside scroll area
        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(12, 12, 12, 12)
        self.extracted_checkboxes_layout.setSpacing(12)

        self.extracted_scroll.setWidget(self.extracted_container)
        self.extracted_section.set_content(self.extracted_scroll)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _add_medication_entry(self):
        """Add a new medication entry row."""
        from PySide6.QtWidgets import QComboBox, QFrame
        from CANONICAL_MEDS import MEDICATIONS

        entry_widget = QFrame()
        entry_widget.setStyleSheet("""
            QFrame {
                background: #f3f4f6;
                border-radius: 8px;
                border: 1px solid #e5e7eb;
            }
        """)
        entry_layout = QVBoxLayout(entry_widget)
        entry_layout.setContentsMargins(10, 8, 10, 8)
        entry_layout.setSpacing(6)

        # Medication name row
        name_row = QHBoxLayout()
        name_row.setSpacing(8)
        name_lbl = QLabel("Med:")
        name_lbl.setStyleSheet("font-size: 17px; color: #374151; font-weight: 500; background: transparent;")
        name_lbl.setFixedWidth(40)
        name_combo = NoWheelComboBox()
        name_combo.setEditable(True)
        name_combo.addItem("")
        name_combo.addItems(sorted(MEDICATIONS.keys()))
        name_combo.setMinimumWidth(180)
        name_combo.setStyleSheet("""
            QComboBox {
                padding: 4px 8px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                background: white;
            }
        """)
        name_row.addWidget(name_lbl)
        name_row.addWidget(name_combo)
        name_row.addStretch()

        # Remove button
        remove_btn = QPushButton("×")
        remove_btn.setFixedSize(24, 24)
        remove_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 4px;
                font-size: 20px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: #dc2626;
            }
        """)
        name_row.addWidget(remove_btn)
        entry_layout.addLayout(name_row)

        # Dose and frequency row
        dose_row = QHBoxLayout()
        dose_row.setSpacing(8)

        dose_lbl = QLabel("Dose:")
        dose_lbl.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500; background: transparent;")
        dose_lbl.setFixedWidth(38)
        dose_combo = NoWheelComboBox()
        dose_combo.setEditable(True)
        dose_combo.setMinimumWidth(100)
        dose_combo.setStyleSheet("""
            QComboBox {
                padding: 4px 8px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                background: white;
            }
        """)
        dose_row.addWidget(dose_lbl)
        dose_row.addWidget(dose_combo)

        freq_lbl = QLabel("Freq:")
        freq_lbl.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500; background: transparent;")
        freq_lbl.setFixedWidth(33)
        freq_combo = NoWheelComboBox()
        freq_combo.addItems(FREQUENCY_OPTIONS)
        freq_combo.setMinimumWidth(100)
        freq_combo.setStyleSheet("""
            QComboBox {
                padding: 4px 8px;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                background: white;
            }
        """)
        dose_row.addWidget(freq_lbl)
        dose_row.addWidget(freq_combo)
        dose_row.addStretch()
        entry_layout.addLayout(dose_row)

        # BNF max info label
        bnf_label = QLabel("")
        bnf_label.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic; background: transparent;")
        entry_layout.addWidget(bnf_label)

        entry_data = {
            "widget": entry_widget,
            "name": name_combo,
            "dose": dose_combo,
            "freq": freq_combo,
            "bnf": bnf_label
        }
        self._medications.append(entry_data)

        def on_med_change(med_name):
            if med_name and med_name in MEDICATIONS:
                info = MEDICATIONS[med_name]
                allowed = info.get("allowed_strengths", [])
                dose_combo.clear()
                if allowed:
                    dose_combo.addItems([f"{s}mg" for s in allowed])
                bnf_max = info.get("bnf_max", "")
                bnf_label.setText(f"Max BNF: {bnf_max}" if bnf_max else "")
            else:
                dose_combo.clear()
                bnf_label.setText("")
            self._update_preview()

        def remove_entry():
            if len(self._medications) > 1:
                self._medications.remove(entry_data)
                entry_widget.deleteLater()
                self._update_preview()

        name_combo.currentTextChanged.connect(on_med_change)
        dose_combo.currentTextChanged.connect(self._update_preview)
        freq_combo.currentIndexChanged.connect(self._update_preview)
        remove_btn.clicked.connect(remove_entry)

        self.med_entries_layout.addWidget(entry_widget)

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        input_parts = []
        extracted_parts = []

        # Build medication text
        med_lines = []
        for entry in self._medications:
            name = entry["name"].currentText().strip()
            dose = entry["dose"].currentText().strip()
            freq = entry["freq"].currentText().strip()
            if name:
                if dose and freq:
                    med_lines.append(f"{name} {dose} {freq}")
                elif dose:
                    med_lines.append(f"{name} {dose}")
                else:
                    med_lines.append(name)

        if med_lines:
            input_parts.append("Current medication: " + ", ".join(med_lines) + ".")

        # Get checked extracted entries
        print(f"[MED POPUP] _update_preview called, {len(self._extracted_checkboxes)} checkboxes")
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                full_text = cb.property("full_text")
                print(f"[MED POPUP] Checkbox checked, full_text={full_text[:50] if full_text else 'None'}...")
                if full_text:
                    extracted_parts.append(full_text)

        # Combine
        if input_parts and extracted_parts:
            combined = "\n\n".join(input_parts) + "\n\n" + "\n\n".join(extracted_parts)
        elif input_parts:
            combined = "\n\n".join(input_parts)
        elif extracted_parts:
            combined = "\n\n".join(extracted_parts)
        else:
            combined = ""

        print(f"[MED POPUP] Emitting combined text, length={len(combined)}")
        # Send directly to card
        self.sent.emit(combined)

    def set_entries(self, entries: list):
        """Set entries from extracted data with dated collapsible cards."""
        from datetime import datetime
        from PySide6.QtWidgets import QTextEdit

        # Clear existing checkboxes
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        # Clear layout
        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Sort entries by date (newest first)
        def get_sort_date(entry):
            if isinstance(entry, dict):
                d = entry.get("date")
                if d:
                    if isinstance(d, datetime):
                        return d
                    return datetime.min
            return datetime.min

        sorted_entries = sorted(entries, key=get_sort_date, reverse=True)

        has_entries = False
        for entry in sorted_entries:
            if isinstance(entry, dict):
                text = entry.get("text", "")
                date = entry.get("date")
                source = entry.get("source", {})
            else:
                text = str(entry)
                date = None
                source = {}

            if not text or not text.strip():
                continue

            has_entries = True

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Get source label - skip Unknown labels
            source_label = source.get("label", "") if isinstance(source, dict) else ""
            if "unknown" in source_label.lower():
                source_label = ""

            # Create entry frame with colored left border
            entry_frame = QFrame()
            entry_frame.setObjectName("medImportEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet("""
                QFrame#medImportEntryFrame {
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-left: 4px solid #008C7E;
                    border-radius: 8px;
                    padding: 2px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 6, 6, 6)
            entry_layout.setSpacing(4)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row: toggle → date → source badge → stretch → checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button on the LEFT
            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(0, 140, 126, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: #008C7E;
                }
                QPushButton:hover { background: rgba(0, 140, 126, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 500;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Source badge
            badge_text = "Medication" if not source_label else source_label[:30]
            source_badge = QLabel(badge_text)
            source_badge.setStyleSheet("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: white;
                    background: #008C7E;
                    border: none;
                    border-radius: 3px;
                    padding: 2px 6px;
                }
            """)
            header_row.addWidget(source_badge)
            header_row.addStretch()

            # Checkbox on the RIGHT
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(self._update_preview)
            header_row.addWidget(cb)
            self._extracted_checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body (full text, hidden by default)
            body_text = QTextEdit()
            body_text.setPlainText(text)
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            body_text.setStyleSheet("""
                QTextEdit {
                    font-size: 16px;
                    color: #333;
                    background: rgba(255, 248, 220, 0.5);
                    border: none;
                    padding: 6px;
                    border-radius: 4px;
                }
            """)
            body_text.setMinimumHeight(50)
            body_text.setMaximumHeight(150)
            body_text.setVisible(False)
            entry_layout.addWidget(body_text)

            drag_bar = QFrame()
            drag_bar.setFixedHeight(8)
            drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
            drag_bar.setStyleSheet("""
                QFrame {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                    border-radius: 2px; margin: 2px 40px;
                }
                QFrame:hover {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                        stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                }
            """)
            drag_bar.setVisible(False)
            drag_bar._drag_y = None
            drag_bar._init_h = None
            def _make_drag_handlers(handle, text_widget):
                def press(ev):
                    handle._drag_y = ev.globalPosition().y()
                    handle._init_h = text_widget.height()
                def move(ev):
                    if handle._drag_y is not None:
                        delta = int(ev.globalPosition().y() - handle._drag_y)
                        new_h = max(60, handle._init_h + delta)
                        text_widget.setMinimumHeight(new_h)
                        text_widget.setMaximumHeight(new_h)
                def release(ev):
                    if handle._drag_y is not None:
                        text_widget.setMaximumHeight(16777215)
                        handle._drag_y = None
                return press, move, release
            dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
            drag_bar.mousePressEvent = dp
            drag_bar.mouseMoveEvent = dm
            drag_bar.mouseReleaseEvent = dr
            entry_layout.addWidget(drag_bar)

            # Toggle function
            def make_toggle(btn, body, frame, popup_self, bar):
                def toggle():
                    is_visible = body.isVisible()
                    body.setVisible(not is_visible)
                    bar.setVisible(not is_visible)
                    btn.setText("▾" if not is_visible else "▸")
                    frame.updateGeometry()
                    if hasattr(popup_self, 'extracted_container'):
                        popup_self.extracted_container.updateGeometry()
                        popup_self.extracted_container.update()
                return toggle

            toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
            toggle_btn.clicked.connect(toggle_fn)
            date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

            self.extracted_checkboxes_layout.addWidget(entry_frame)

        self.extracted_checkboxes_layout.addStretch()

        if has_entries:
            self.extracted_section.setVisible(True)
            if self.extracted_section.is_collapsed():
                self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)

    def get_state(self) -> dict:
        """Get current state for saving."""
        meds = []
        for entry in self._medications:
            name = entry["name"].currentText().strip()
            dose = entry["dose"].currentText().strip()
            freq = entry["freq"].currentText().strip()
            if name:
                meds.append({"name": name, "dose": dose, "freq": freq})
        return {"medications": meds}

    def load_state(self, state: dict):
        """Load state from saved data."""
        meds = state.get("medications", [])
        # Clear existing entries
        while len(self._medications) > 1:
            entry = self._medications.pop()
            entry["widget"].deleteLater()

        for i, med in enumerate(meds):
            if i >= len(self._medications):
                self._add_medication_entry()
            entry = self._medications[i]
            entry["name"].setCurrentText(med.get("name", ""))
            entry["dose"].setCurrentText(med.get("dose", ""))
            freq = med.get("freq", "OD")
            idx = entry["freq"].findText(freq)
            if idx >= 0:
                entry["freq"].setCurrentIndex(idx)

        self._update_preview()

    def prefill_medications_from_notes(self, raw_notes: list):
        """Extract medications from notes (last 6 months) and pre-fill current medication section.

        Finds most recent medication per class, prioritizing psychiatric medications.
        """
        import re
        from datetime import datetime, timedelta
        from CANONICAL_MEDS import MEDICATIONS

        if not raw_notes:
            print(f"[GPR-MED] No raw notes for medication extraction")
            return

        # === PARSE DATE HELPER ===
        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # === FIND MOST RECENT NOTE DATE AND FILTER TO LAST 6 MONTHS ===
        note_dates = [parse_date(n.get("date") or n.get("datetime")) for n in raw_notes]
        note_dates = [d for d in note_dates if d]
        if not note_dates:
            print(f"[GPR-MED] No dates found in notes")
            return

        latest_date = max(note_dates)
        cutoff_date = latest_date - timedelta(days=180)  # 6 months
        print(f"[GPR-MED] Latest note: {latest_date.date()}, searching from {cutoff_date.date()} (last 6 months)")

        # Filter notes to last 6 months only
        recent_notes = []
        for n in raw_notes:
            note_date = parse_date(n.get("date") or n.get("datetime"))
            if note_date and note_date >= cutoff_date:
                recent_notes.append(n)

        print(f"[GPR-MED] Extracting medications from {len(recent_notes)} notes (last 6 months)...")

        # === BUILD TOKEN INDEX ===
        token_map = {}  # token -> (key, canonical)
        meta_map = {}   # key -> metadata

        for key, meta in MEDICATIONS.items():
            meta_map[key] = meta
            canonical = meta["canonical"]
            for syn in meta.get("patterns", []):
                s = syn.lower().strip()
                if s:
                    token_map[s] = (key, canonical)

        def tokenise(text):
            text = text.lower()
            text = re.sub(r'(?<=\d)\.(?=\d)', 'DOT', text)
            text = re.sub(r'[^a-z0-9DOT]+', ' ', text)
            text = text.replace("DOT", ".")
            return text.split()

        def find_dose(tokens, idx):
            unit_set = {"mg", "mcg", "µg", "g", "units", "iu"}
            for i in range(idx + 1, min(idx + 4, len(tokens))):
                tok = tokens[i]
                m = re.match(r'(\d+(?:\.\d+)?)(mg|mcg|µg|g|units|iu)$', tok)
                if m:
                    return float(m.group(1)), m.group(2)
                if tok.isdigit() or re.match(r'\d+\.\d+', tok):
                    if i + 1 < len(tokens) and tokens[i + 1] in unit_set:
                        return float(tok), tokens[i + 1]
            return None, None

        def find_freq(tokens, idx):
            freq_set = {"od", "bd", "tds", "qds", "qid", "nocte", "mane", "stat", "prn", "daily", "weekly", "monthly"}
            for i in range(idx, min(idx + 5, len(tokens))):
                if tokens[i] in freq_set:
                    return tokens[i]
            return None

        # === PSYCHIATRIC MEDICATION CLASSES (prioritized) ===
        PSYCH_CLASSES = {
            "Antipsychotic": 1,
            "Antidepressant": 2,
            "Mood Stabiliser": 3,
            "Anxiolytic": 4,
            "Benzodiazepine": 5,
            "Anticonvulsant / Benzodiazepine": 5,
            "Sedative": 6,
            "Sedation": 6,
            "Sleep": 7,
            "Stimulant": 8,
            "Addictions": 9,
            "Opioid Substitution": 9,
        }

        # === EXTRACT MEDICATIONS ===
        meds_found = []
        for n in recent_notes:
            content = n.get("content", "") or n.get("text", "") or ""
            if not content:
                continue
            tokens = tokenise(content)
            note_date = parse_date(n.get("date") or n.get("datetime"))

            for i, tok in enumerate(tokens):
                if tok in token_map:
                    key, canonical = token_map[tok]
                    meta = meta_map[key]
                    med_class = meta.get("class", "Other")

                    strength, unit = find_dose(tokens, i)
                    freq = find_freq(tokens, i)

                    meds_found.append({
                        "med_key": key,
                        "canonical": canonical,
                        "class": med_class,
                        "strength": strength,
                        "unit": unit or "mg",
                        "frequency": freq,
                        "date": note_date,
                    })

        print(f"[GPR-MED] Found {len(meds_found)} medication mentions in last 6 months")

        if not meds_found:
            print(f"[GPR-MED] No medications found in notes")
            return

        # === GROUP BY CLASS AND PICK MOST RECENT PER CLASS ===
        class_groups = {}
        for med in meds_found:
            med_class = med.get("class", "Other")
            if med_class not in class_groups:
                class_groups[med_class] = []
            class_groups[med_class].append(med)

        # For each class, find the most recent medication with a dose
        most_recent_by_class = []
        for med_class, mentions in class_groups.items():
            # Sort by date descending
            sorted_mentions = sorted(
                mentions,
                key=lambda x: x.get("date") or datetime.min,
                reverse=True
            )
            # Take most recent with a dose
            for m in sorted_mentions:
                if m.get("strength"):
                    most_recent_by_class.append(m)
                    break
            else:
                # If none have dose, take most recent anyway
                if sorted_mentions:
                    most_recent_by_class.append(sorted_mentions[0])

        # === PRIORITIZE PSYCHIATRIC CLASSES ===
        def class_priority(med):
            med_class = med.get("class", "Other")
            return PSYCH_CLASSES.get(med_class, 100)  # Non-psych classes get low priority

        # Sort: psych classes first (by priority), then others
        most_recent_by_class.sort(key=class_priority)

        # Limit to top entries (psych meds first)
        max_meds = 10
        final_meds = most_recent_by_class[:max_meds]

        print(f"[GPR-MED] Selected {len(final_meds)} medications (1 per class, psych prioritized):")
        for m in final_meds:
            strength = m.get('strength')
            dose_str = f"{strength}{m.get('unit', 'mg')}" if strength else "no dose"
            print(f"[GPR-MED]   [{m.get('class')}] {m.get('canonical')}: {dose_str} {m.get('frequency') or ''} ({m.get('date').date() if m.get('date') else 'no date'})")

        # Map extracted frequency to popup frequency options
        FREQ_MAP = {
            "od": "OD", "daily": "OD", "once daily": "OD", "1x daily": "OD",
            "bd": "BD", "twice daily": "BD", "2x daily": "BD",
            "tds": "TDS", "three times daily": "TDS", "3x daily": "TDS",
            "qds": "QDS", "qid": "QDS", "four times daily": "QDS", "4x daily": "QDS",
            "nocte": "Nocte", "at night": "Nocte", "night": "Nocte", "on": "Nocte",
            "mane": "Mane", "in the morning": "Mane", "am": "Mane",
            "prn": "PRN", "as required": "PRN", "when required": "PRN",
            "weekly": "Weekly", "1 weekly": "Weekly", "once weekly": "Weekly",
            "fortnightly": "Fortnightly", "2 weekly": "Fortnightly", "every 2 weeks": "Fortnightly",
            "3 weekly": "3 Weekly", "every 3 weeks": "3 Weekly",
            "monthly": "Monthly", "4 weekly": "Monthly", "every 4 weeks": "Monthly",
        }

        # Populate the medication entries
        while len(self._medications) < len(final_meds):
            self._add_medication_entry()

        for i, med in enumerate(final_meds):
            if i >= len(self._medications):
                break

            entry = self._medications[i]
            med_key = med.get("med_key", "")

            # Set medication name (the key in MEDICATIONS dict)
            name_combo = entry.get("name")
            if name_combo:
                idx = name_combo.findText(med_key)
                if idx >= 0:
                    name_combo.setCurrentIndex(idx)
                else:
                    # Try canonical name
                    canonical = med.get("canonical", "")
                    idx = name_combo.findText(canonical.upper())
                    if idx >= 0:
                        name_combo.setCurrentIndex(idx)

            # Set dose
            dose_combo = entry.get("dose")
            if dose_combo and med.get("strength"):
                strength = med.get("strength")
                unit = med.get("unit", "mg")
                if isinstance(strength, float) and strength == int(strength):
                    dose_str = f"{int(strength)}{unit}"
                else:
                    dose_str = f"{strength}{unit}"
                idx = dose_combo.findText(dose_str)
                if idx >= 0:
                    dose_combo.setCurrentIndex(idx)
                else:
                    dose_combo.setCurrentText(dose_str)

            # Set frequency
            freq_combo = entry.get("freq")
            if freq_combo and med.get("frequency"):
                freq = med.get("frequency", "").lower()
                mapped_freq = FREQ_MAP.get(freq, "OD")
                idx = freq_combo.findText(mapped_freq)
                if idx >= 0:
                    freq_combo.setCurrentIndex(idx)

        # Update preview
        self._update_preview()

        print(f"[GPR-MED] ✓ Pre-filled {len(final_meds)} medication(s)")


# ================================================================
# GPR DIAGNOSIS POPUP (Mental Disorder - ICD-10)
# ================================================================

class GPRDiagnosisPopup(QWidget):
    """Diagnosis popup with preview, ICD-10 selection, and extracted data sections."""

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._extracted_checkboxes = []
        self._load_icd10_data()
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _load_icd10_data(self):
        """Load ICD-10 diagnosis dictionary."""
        try:
            from icd10_dict import ICD10_DICT
            self.icd10_dict = ICD10_DICT
        except ImportError:
            self.icd10_dict = {}

    def _setup_ui(self):
        from PySide6.QtWidgets import QComboBox, QCompleter, QStyleFactory

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(4)

        # Main scroll area for inputs and imported data (scrollable)
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: DIAGNOSIS INPUT (collapsible with drag bar)
        # ====================================================
        self.input_section = CollapsibleSection("Diagnosis (ICD-10)", start_collapsed=True)
        self.input_section.set_content_height(280)
        self.input_section._min_height = 100
        self.input_section._max_height = 400
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.input_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #005a50;
                background: transparent;
                border: none;
            }
        """)

        # Diagnosis form container
        dx_form_scroll = QScrollArea()
        dx_form_scroll.setWidgetResizable(True)
        dx_form_scroll.setFrameShape(QScrollArea.NoFrame)
        dx_form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        dx_form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        dx_form_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255,255,255,0.95);
                border: 1px solid rgba(0, 140, 126, 0.2);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        dx_form_container = QWidget()
        dx_form_container.setStyleSheet("background: transparent;")
        dx_form_layout = QVBoxLayout(dx_form_container)
        dx_form_layout.setContentsMargins(12, 12, 12, 12)
        dx_form_layout.setSpacing(12)

        # Diagnosis comboboxes (up to 3)
        self.dx_boxes = []

        for i in range(3):
            # Label for each diagnosis slot
            slot_label = QLabel(f"Diagnosis {i + 1}:")
            slot_label.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151; background: transparent;")
            dx_form_layout.addWidget(slot_label)

            combo = NoWheelComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.setInsertPolicy(NoWheelComboBox.InsertPolicy.NoInsert)
            combo.lineEdit().setPlaceholderText("Start typing to search...")

            combo.setSizeAdjustPolicy(QComboBox.SizeAdjustPolicy.AdjustToMinimumContentsLengthWithIcon)
            combo.setMinimumContentsLength(25)
            combo.setStyleSheet("""
                QComboBox {
                    padding: 8px;
                    font-size: 17px;
                    border: 1px solid #d1d5db;
                    border-radius: 6px;
                    background-color: #ffffff;
                }
                QComboBox QLineEdit {
                    background-color: #ffffff;
                    border: none;
                    padding: 0px;
                }
                QComboBox::drop-down {
                    border: none;
                    background-color: #f3f4f6;
                    width: 24px;
                    border-top-right-radius: 5px;
                    border-bottom-right-radius: 5px;
                }
                QComboBox QAbstractItemView {
                    min-width: 400px;
                    background-color: #ffffff;
                    selection-background-color: #e5e7eb;
                }
            """)

            # Add items
            combo.addItem("Not specified", None)

            for diagnosis, meta in sorted(
                self.icd10_dict.items(),
                key=lambda x: x[0].lower()
            ):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(
                    diagnosis,
                    {"diagnosis": diagnosis, "icd10": icd_code}
                )

            # Autocomplete
            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)

            combo.setMaxVisibleItems(15)
            combo.currentIndexChanged.connect(self._update_preview)

            dx_form_layout.addWidget(combo)
            self.dx_boxes.append(combo)

        dx_form_layout.addStretch()
        dx_form_scroll.setWidget(dx_form_container)
        self.input_section.set_content(dx_form_scroll)
        main_layout.addWidget(self.input_section)

        # ====================================================
        # SECTION 3: IMPORTED DATA (collapsible, at bottom)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 400
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        # Use QScrollArea directly as content
        self.extracted_scroll = QScrollArea()
        self.extracted_scroll.setWidgetResizable(True)
        self.extracted_scroll.setFrameShape(QScrollArea.NoFrame)
        self.extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.extracted_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        # Container for checkboxes inside scroll area
        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(12, 12, 12, 12)
        self.extracted_checkboxes_layout.setSpacing(12)

        self.extracted_scroll.setWidget(self.extracted_container)
        self.extracted_section.set_content(self.extracted_scroll)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        input_parts = []
        extracted_parts = []

        # Build diagnosis text
        diagnoses = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    if icd:
                        diagnoses.append(f"{dx} ({icd})")
                    else:
                        diagnoses.append(dx)

        if diagnoses:
            if len(diagnoses) == 1:
                input_parts.append(f"{diagnoses[0]} is a mental disorder as defined by the Mental Health Act.")
            else:
                joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                input_parts.append(f"{joined} are mental disorders as defined by the Mental Health Act.")

        # Get checked extracted entries
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                extracted_parts.append(cb.property("full_text"))

        # Combine
        if input_parts and extracted_parts:
            combined = "\n\n".join(input_parts) + "\n\n" + "\n\n".join(extracted_parts)
        elif input_parts:
            combined = "\n\n".join(input_parts)
        elif extracted_parts:
            combined = "\n\n".join(extracted_parts)
        else:
            combined = ""

        # Send directly to card
        self.sent.emit(combined)

    def set_entries(self, entries: list):
        """Set entries from extracted data with dated collapsible cards."""
        from datetime import datetime
        from PySide6.QtWidgets import QTextEdit

        # Clear existing checkboxes
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        # Clear layout
        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Sort entries by date (newest first)
        def get_sort_date(entry):
            if isinstance(entry, dict):
                d = entry.get("date")
                if d and isinstance(d, datetime):
                    return d
            return datetime.min

        sorted_entries = sorted(entries, key=get_sort_date, reverse=True)

        has_entries = False
        for entry in sorted_entries:
            if isinstance(entry, dict):
                text = entry.get("text", "")
                date = entry.get("date")
                source = entry.get("source", {})
            else:
                text = str(entry)
                date = None
                source = {}

            if not text or not text.strip():
                continue

            has_entries = True

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Get source label - skip Unknown labels
            source_label = source.get("label", "") if isinstance(source, dict) else ""
            if "unknown" in source_label.lower():
                source_label = ""

            # Create entry frame with colored left border
            entry_frame = QFrame()
            entry_frame.setObjectName("dxImportEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet("""
                QFrame#dxImportEntryFrame {
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-left: 4px solid #9c27b0;
                    border-radius: 8px;
                    padding: 2px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 6, 6, 6)
            entry_layout.setSpacing(4)

            # Header row: toggle → date → source badge → stretch → checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button on the LEFT
            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(156, 39, 176, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: #9c27b0;
                }
                QPushButton:hover { background: rgba(156, 39, 176, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 500;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Source badge
            badge_text = "Diagnosis" if not source_label else source_label[:30]
            source_badge = QLabel(badge_text)
            source_badge.setStyleSheet("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: white;
                    background: #9c27b0;
                    border: none;
                    border-radius: 3px;
                    padding: 2px 6px;
                }
            """)
            header_row.addWidget(source_badge)
            header_row.addStretch()

            # Checkbox on the RIGHT
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(self._update_preview)
            header_row.addWidget(cb)
            self._extracted_checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body (full text, hidden by default) - with diagnosis highlights
            body_text = QTextEdit()
            highlighted_html = self._highlight_diagnosis_patterns(text)
            body_text.setHtml(highlighted_html)
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            body_text.setStyleSheet("""
                QTextEdit {
                    font-size: 16px;
                    color: #333;
                    background: rgba(255, 248, 220, 0.5);
                    border: none;
                    padding: 6px;
                    border-radius: 4px;
                }
            """)
            body_text.setMinimumHeight(50)
            body_text.setMaximumHeight(150)
            body_text.setVisible(False)
            entry_layout.addWidget(body_text)

            # Toggle function
            def make_toggle(btn, body, frame, popup_self):
                def toggle():
                    is_visible = body.isVisible()
                    body.setVisible(not is_visible)
                    btn.setText("▾" if not is_visible else "▸")
                    frame.updateGeometry()
                    if hasattr(popup_self, 'extracted_container'):
                        popup_self.extracted_container.updateGeometry()
                        popup_self.extracted_container.update()
                return toggle

            toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self)
            toggle_btn.clicked.connect(toggle_fn)
            date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

            self.extracted_checkboxes_layout.addWidget(entry_frame)

        if has_entries:
            self.extracted_section.setVisible(True)
            if self.extracted_section.is_collapsed():
                self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)

    def _has_diagnosis_evidence(self, text: str) -> bool:
        """Check if text contains any diagnosis-related patterns."""
        import re
        text_lower = text.lower()

        # Diagnosis patterns to check for
        DIAGNOSIS_PATTERNS = [
            r'paranoid\s+schizophrenia', r'catatonic\s+schizophrenia', r'hebephrenic\s+schizophrenia',
            r'residual\s+schizophrenia', r'simple\s+schizophrenia', r'undifferentiated\s+schizophrenia',
            r'schizoaffective', r'schizophrenia',
            r'atypical\s+autism', r"asperger['\"]?s?\s+syndrome", r'diagnosed\s+with\s+asperger',
            r'childhood\s+autism', r'autism\s+spectrum\s+disorder', r'autistic\s+spectrum\s+disorder', r'\basd\b',
            r'bipolar\s+affective\s+disorder', r'bipolar\s+disorder', r'manic\s+depressi',
            r'recurrent\s+depressi', r'major\s+depressi', r'depressi(?:ve|on)',
            r'emotionally\s+unstable\s+personality', r'\beupd\b', r'borderline\s+personality',
            r'antisocial\s+personality', r'dissocial\s+personality', r'narcissistic\s+personality',
            r'paranoid\s+personality', r'personality\s+disorder',
            r'generalised\s+anxiety', r'generalized\s+anxiety', r'\bptsd\b', r'post[- ]?traumatic\s+stress',
            r'acute.*psycho(?:tic|sis)',
            r'learning\s+disabilit', r'intellectual\s+disabilit',
            r'alcohol\s+dependence', r'drug\s+dependence', r'opioid\s+dependence',
            r'substance\s+misuse', r'substance\s+abuse',
        ]

        for pattern in DIAGNOSIS_PATTERNS:
            if re.search(pattern, text_lower):
                return True
        return False

    def _highlight_diagnosis_patterns(self, text: str) -> str:
        """Highlight diagnosis-related patterns in text using HTML."""
        import re
        import html

        # Escape HTML special characters first
        escaped_text = html.escape(text)

        # Diagnosis patterns to highlight (regex pattern, highlight color)
        HIGHLIGHT_PATTERNS = [
            # Schizophrenia variants
            (r'(?i)(paranoid\s+schizophrenia)', '#ffeb3b'),  # Yellow
            (r'(?i)(catatonic\s+schizophrenia)', '#ffeb3b'),
            (r'(?i)(hebephrenic\s+schizophrenia)', '#ffeb3b'),
            (r'(?i)(residual\s+schizophrenia)', '#ffeb3b'),
            (r'(?i)(simple\s+schizophrenia)', '#ffeb3b'),
            (r'(?i)(undifferentiated\s+schizophrenia)', '#ffeb3b'),
            (r'(?i)(schizoaffective)', '#ffeb3b'),
            (r'(?i)(schizophrenia)', '#ffeb3b'),
            # Autism
            (r'(?i)(atypical\s+autism)', '#b3e5fc'),  # Light blue
            (r"(?i)(asperger['\"]?s?\s+syndrome)", '#b3e5fc'),
            (r'(?i)(childhood\s+autism)', '#b3e5fc'),
            (r'(?i)(autism\s+spectrum\s+disorder)', '#b3e5fc'),
            (r'(?i)(autistic\s+spectrum\s+disorder)', '#b3e5fc'),
            (r'(?i)\b(asd)\b', '#b3e5fc'),
            # Mood disorders
            (r'(?i)(bipolar\s+affective\s+disorder)', '#c8e6c9'),  # Light green
            (r'(?i)(bipolar\s+disorder)', '#c8e6c9'),
            (r'(?i)(manic\s+depressi\w*)', '#c8e6c9'),
            (r'(?i)(recurrent\s+depressi\w*)', '#c8e6c9'),
            (r'(?i)(major\s+depressi\w*)', '#c8e6c9'),
            (r'(?i)(depressi(?:ve|on))', '#c8e6c9'),
            # Personality disorders
            (r'(?i)(emotionally\s+unstable\s+personality)', '#f8bbd9'),  # Light pink
            (r'(?i)\b(eupd)\b', '#f8bbd9'),
            (r'(?i)(borderline\s+personality)', '#f8bbd9'),
            (r'(?i)(antisocial\s+personality)', '#f8bbd9'),
            (r'(?i)(dissocial\s+personality)', '#f8bbd9'),
            (r'(?i)(narcissistic\s+personality)', '#f8bbd9'),
            (r'(?i)(paranoid\s+personality)', '#f8bbd9'),
            (r'(?i)(personality\s+disorder)', '#f8bbd9'),
            # Anxiety/PTSD
            (r'(?i)(generalised\s+anxiety)', '#fff9c4'),  # Light yellow
            (r'(?i)(generalized\s+anxiety)', '#fff9c4'),
            (r'(?i)\b(ptsd)\b', '#fff9c4'),
            (r'(?i)(post[- ]?traumatic\s+stress)', '#fff9c4'),
            # Psychosis
            (r'(?i)(acute.*?psycho(?:tic|sis))', '#e1bee7'),  # Light purple
            # Learning disability
            (r'(?i)(learning\s+disabilit\w*)', '#d7ccc8'),  # Light brown
            (r'(?i)(intellectual\s+disabilit\w*)', '#d7ccc8'),
            # Substance
            (r'(?i)(alcohol\s+dependence)', '#ffccbc'),  # Light orange
            (r'(?i)(drug\s+dependence)', '#ffccbc'),
            (r'(?i)(opioid\s+dependence)', '#ffccbc'),
            (r'(?i)(substance\s+misuse)', '#ffccbc'),
            (r'(?i)(substance\s+abuse)', '#ffccbc'),
        ]

        result = escaped_text
        already_matched = set()

        for pattern, color in HIGHLIGHT_PATTERNS:
            def replace_match(m):
                matched_text = m.group(1)
                # Avoid double-highlighting
                if matched_text.lower() in already_matched:
                    return matched_text
                already_matched.add(matched_text.lower())
                return f'<span style="background-color: {color}; padding: 1px 3px; border-radius: 2px;">{matched_text}</span>'

            result = re.sub(pattern, replace_match, result)

        # Wrap in basic HTML structure with styling
        html_result = f'''<html>
<body style="font-family: -apple-system, BlinkMacSystemFont, sans-serif; font-size: 16px; color: #333; line-height: 1.4;">
{result.replace(chr(10), '<br>')}
</body>
</html>'''

        return html_result

    def prefill_diagnoses_from_notes(self, raw_notes: list, imported_entries: list = None, substance_entries: list = None):
        """Extract diagnoses from imported data (precedence) then notes.

        Imported data takes precedence as it contains validated historical diagnoses.
        Only 1 schizophrenia-type diagnosis is allowed (schizophrenia OR schizoaffective).
        For substance misuse, checks section 8 data to identify specific drug and recency.
        """
        import re
        from datetime import datetime, timedelta

        extracted_diagnoses = []
        matched_categories = set()
        substance_diagnosis_needed = False  # Flag to handle substance after main extraction

        def find_icd10_entry(search_term, course_preference="continuous"):
            """Find ICD-10 entry, preferring specified course type."""
            search_lower = search_term.lower()
            best_match = None
            preferred_match = None

            for diag_name, meta in self.icd10_dict.items():
                diag_lower = diag_name.lower()
                if search_lower in diag_lower:
                    # Check for preferred course type
                    if course_preference == "continuous" and "continuous" in diag_lower:
                        preferred_match = diag_name
                        break
                    elif course_preference == "remission" and "remission" in diag_lower:
                        preferred_match = diag_name
                        break
                    elif best_match is None:
                        best_match = diag_name

            return preferred_match or best_match

        def detect_substance_from_section8(substance_entries):
            """Analyze substance_use entries (section 8) to find specific drug and recency.

            Returns: (drug_name, is_current) where drug_name is the primary drug found
            and is_current indicates use within 3 months of latest entry.
            """
            if not substance_entries:
                return None, False

            # Drug patterns mapped to ICD-10 search terms
            DRUG_PATTERNS = [
                (r'\bcocaine\b', 'cocaine'),
                (r'\bcrack\b', 'cocaine'),
                (r'\bheroin\b', 'opioids'),
                (r'\bopioid', 'opioids'),
                (r'\bopiate', 'opioids'),
                (r'\bmorphine\b', 'opioids'),
                (r'\bfentanyl\b', 'opioids'),
                (r'\bcannabis\b', 'cannabinoids'),
                (r'\bmarijuana\b', 'cannabinoids'),
                (r'\bweed\b', 'cannabinoids'),
                (r'\bskunk\b', 'cannabinoids'),
                (r'\balcohol\b', 'alcohol'),
                (r'\bamphetamine', 'stimulants'),
                (r'\bmeth\b', 'stimulants'),
                (r'\bspeed\b', 'stimulants'),
                (r'\bmdma\b', 'stimulants'),
                (r'\becstasy\b', 'stimulants'),
                (r'\bbenzodiazepin', 'sedatives'),
                (r'\bvalium\b', 'sedatives'),
                (r'\bdiazepam\b', 'sedatives'),
                (r'\bxanax\b', 'sedatives'),
            ]

            # ICD-10 M&BD terms for each drug type (matching ICD10_DICT.txt format)
            DRUG_TO_ICD10 = {
                'cocaine': 'M&BD - cocaine active dependence',
                'opioids': 'M&BD - opioids active dependence',
                'cannabinoids': 'M&BD - cannabinoids active dependence',
                'alcohol': 'M&BD - alcohol active dependence',
                'stimulants': 'M&BD - other stimulants',
                'sedatives': 'M&BD - sedatives or hypnotics active dependence',
            }

            # Count drug mentions and track dates
            drug_counts = {}
            drug_dates = {}
            latest_date = None

            for entry in substance_entries:
                text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
                text_lower = text.lower()
                entry_date = entry.get("date") if isinstance(entry, dict) else None

                # Track latest date
                if isinstance(entry_date, datetime):
                    if latest_date is None or entry_date > latest_date:
                        latest_date = entry_date

                # Find drug mentions
                for pattern, drug_type in DRUG_PATTERNS:
                    if re.search(pattern, text_lower):
                        drug_counts[drug_type] = drug_counts.get(drug_type, 0) + 1
                        if drug_type not in drug_dates:
                            drug_dates[drug_type] = []
                        if isinstance(entry_date, datetime):
                            drug_dates[drug_type].append(entry_date)

            if not drug_counts:
                print("[GPR-DX] No specific drugs found in substance_use entries")
                return None, False

            # Find primary drug (most mentioned)
            primary_drug = max(drug_counts, key=drug_counts.get)
            print(f"[GPR-DX] Drug counts: {drug_counts}, primary: {primary_drug}")

            # Check if current (use within 3 months of latest entry)
            is_current = False
            if latest_date and primary_drug in drug_dates:
                three_months_ago = latest_date - timedelta(days=90)
                recent_uses = [d for d in drug_dates[primary_drug] if d >= three_months_ago]
                is_current = len(recent_uses) > 0
                print(f"[GPR-DX] Latest date: {latest_date}, recent uses: {len(recent_uses)}, is_current: {is_current}")

            icd10_term = DRUG_TO_ICD10.get(primary_drug)
            return icd10_term, is_current

        # Diagnosis patterns - schizophrenia and schizoaffective share same category
        # NOTE: Substance patterns are now handled separately via detect_substance_from_section8
        DIAGNOSIS_PATTERNS = [
            # Schizophrenia variants (all share 'schizo' category - only 1 allowed)
            (r'paranoid\s+schizophrenia', 'Paranoid schizophrenia', 'schizo'),
            (r'catatonic\s+schizophrenia', 'Catatonic schizophrenia', 'schizo'),
            (r'hebephrenic\s+schizophrenia', 'Hebephrenic schizophrenia', 'schizo'),
            (r'residual\s+schizophrenia', 'Residual schizophrenia', 'schizo'),
            (r'simple\s+schizophrenia', 'Simple schizophrenia', 'schizo'),
            (r'undifferentiated\s+schizophrenia', 'Undifferentiated schizophrenia', 'schizo'),
            (r'schizoaffective', 'Schizoaffective disorder', 'schizo'),
            (r'schizophrenia', 'Schizophrenia, unspecified', 'schizo'),
            # Autism - only auto-detect if explicitly stated (removed loose patterns to avoid false positives)
            (r'atypical\s+autism', 'Atypical autism', 'autism'),
            (r"asperger['\"]?s?\s+syndrome", 'Asperger', 'autism'),
            (r'diagnosed\s+with\s+asperger', 'Asperger', 'autism'),
            (r'childhood\s+autism', 'Childhood autism', 'autism'),
            (r'autism\s+spectrum\s+disorder', 'Childhood autism', 'autism'),
            (r'autistic\s+spectrum\s+disorder', 'Childhood autism', 'autism'),
            (r'\basd\b', 'Childhood autism', 'autism'),
            # Mood disorders
            (r'bipolar\s+affective\s+disorder', 'Bipolar affective disorder', 'bipolar'),
            (r'bipolar\s+disorder', 'Bipolar affective disorder', 'bipolar'),
            (r'manic\s+depressi', 'Bipolar affective disorder', 'bipolar'),
            (r'recurrent\s+depressi', 'Recurrent depressive disorder', 'depression'),
            (r'major\s+depressi', 'Depressive episode', 'depression'),
            (r'depressi(?:ve|on)', 'Depressive episode', 'depression'),
            # Personality disorders
            (r'emotionally\s+unstable\s+personality', 'Emotionally unstable personality disorder', 'personality'),
            (r'eupd', 'Emotionally unstable personality disorder', 'personality'),
            (r'borderline\s+personality', 'Emotionally unstable personality disorder', 'personality'),
            (r'antisocial\s+personality', 'Dissocial personality disorder', 'personality'),
            (r'dissocial\s+personality', 'Dissocial personality disorder', 'personality'),
            (r'narcissistic\s+personality', 'Other specific personality disorders', 'personality'),
            (r'paranoid\s+personality', 'Paranoid personality disorder', 'personality'),
            (r'personality\s+disorder', 'Personality disorder', 'personality'),
            # Anxiety
            (r'generalised\s+anxiety', 'Generalized anxiety disorder', 'anxiety'),
            (r'generalized\s+anxiety', 'Generalized anxiety disorder', 'anxiety'),
            (r'ptsd', 'Post-traumatic stress disorder', 'ptsd'),
            (r'post[- ]?traumatic\s+stress', 'Post-traumatic stress disorder', 'ptsd'),
            # Psychosis
            (r'acute.*psycho(?:tic|sis)', 'Acute and transient psychotic disorders', 'psychosis'),
            # Learning disability
            (r'learning\s+disabilit', 'Mental retardation', 'learning'),
            (r'intellectual\s+disabilit', 'Mental retardation', 'learning'),
            # Substance - flag for intelligent detection via section 8
            (r'substance\s+misuse', '__SUBSTANCE_FLAG__', 'substance'),
            (r'substance\s+abuse', '__SUBSTANCE_FLAG__', 'substance'),
            (r'polysubstance', '__SUBSTANCE_FLAG__', 'substance'),
            (r'poly[\s-]?drug', '__SUBSTANCE_FLAG__', 'substance'),
            (r'multiple\s+drug\s+use', '__SUBSTANCE_FLAG__', 'substance'),
            (r'alcohol\s+dependence', 'Alcohol dependence syndrome', 'alcohol'),
            (r'drug\s+dependence', 'Drug dependence', 'drugs'),
            (r'opioid\s+dependence', 'Opioid dependence', 'drugs'),
        ]

        def extract_from_text(text, source_name):
            """Extract diagnoses from text, respecting category limits."""
            nonlocal extracted_diagnoses, matched_categories, substance_diagnosis_needed
            text_lower = text.lower()

            # Check for schizophrenia course indicators in this text
            def get_schizo_course():
                if "complete remission" in text_lower or "in remission" in text_lower:
                    return "remission"
                elif "episodic" in text_lower:
                    return "episodic"
                return "continuous"

            schizo_course = get_schizo_course()

            # Negation phrases to check before confirming a diagnosis
            NEGATION_PHRASES = [
                "no ", "not ", "ruled out", "no evidence", "unlikely", "excluded",
                "does not have", "doesn't have", "without ", "denies ", "negative for"
            ]

            def is_negated(text, match_pos):
                """Check if a match is negated by looking at preceding text."""
                # Look at 50 chars before the match
                start = max(0, match_pos - 50)
                preceding = text[start:match_pos].lower()
                for neg in NEGATION_PHRASES:
                    if neg in preceding:
                        return True
                return False

            for pattern, search_term, category in DIAGNOSIS_PATTERNS:
                if category in matched_categories:
                    continue
                if len(extracted_diagnoses) >= 3:
                    break

                match = re.search(pattern, text_lower)
                if match:
                    # Check if this match is negated
                    if is_negated(text_lower, match.start()):
                        print(f"[GPR-DX] {source_name}: Skipping negated '{pattern}'")
                        continue

                    # Handle substance misuse flag - defer to intelligent detection
                    if search_term == '__SUBSTANCE_FLAG__':
                        substance_diagnosis_needed = True
                        matched_categories.add(category)
                        print(f"[GPR-DX] {source_name}: '{pattern}' -> flagged for intelligent substance detection")
                        continue

                    # Handle schizophrenia course
                    if category == 'schizo':
                        display = find_icd10_entry(search_term, schizo_course)
                    else:
                        display = find_icd10_entry(search_term)

                    if display and display not in extracted_diagnoses:
                        extracted_diagnoses.append(display)
                        matched_categories.add(category)
                        print(f"[GPR-DX] {source_name}: '{pattern}' -> '{display}'")

        # === STEP 1: Check imported data FIRST (takes precedence) ===
        if imported_entries:
            print(f"[GPR-DX] Checking {len(imported_entries)} imported entries (precedence)")
            # Sort by date (newest first) for most recent diagnoses
            def get_date(e):
                if isinstance(e, dict):
                    d = e.get("date")
                    if isinstance(d, datetime):
                        return d
                return datetime.min

            sorted_imports = sorted(imported_entries, key=get_date, reverse=True)

            for entry in sorted_imports:
                if len(extracted_diagnoses) >= 3:
                    break
                text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
                if text:
                    extract_from_text(text, "IMPORTED")

        # === STEP 2: Handle substance misuse with intelligent detection (BEFORE notes) ===
        # This ensures substance gets priority over random matches in notes
        if substance_diagnosis_needed and len(extracted_diagnoses) < 3:
            print(f"[GPR-DX] Substance misuse detected - checking section 8 for specific drug")
            if substance_entries:
                print(f"[GPR-DX] Analyzing {len(substance_entries)} substance_use entries")
                drug_icd10_term, is_current = detect_substance_from_section8(substance_entries)

                if drug_icd10_term:
                    # Find the ICD-10 entry for this drug
                    display = find_icd10_entry(drug_icd10_term)
                    print(f"[GPR-DX] SUBSTANCE search: '{drug_icd10_term}' -> found: '{display}'")
                    if not display:
                        # Try shorter search - just the drug name
                        drug_name = drug_icd10_term.split("due to use of ")[-1] if "due to use of" in drug_icd10_term else drug_icd10_term
                        print(f"[GPR-DX] SUBSTANCE retry with: '{drug_name}'")
                        display = find_icd10_entry(drug_name)
                        print(f"[GPR-DX] SUBSTANCE retry result: '{display}'")
                    if display and display not in extracted_diagnoses:
                        extracted_diagnoses.append(display)
                        status = "current use" if is_current else "historical"
                        print(f"[GPR-DX] SUBSTANCE: '{drug_icd10_term}' -> '{display}' ({status})")
                else:
                    # Fallback to generic multiple drug use
                    display = find_icd10_entry("Mental and behavioural disorders due to multiple drug use")
                    if display and display not in extracted_diagnoses:
                        extracted_diagnoses.append(display)
                        print(f"[GPR-DX] SUBSTANCE: No specific drug found, using multiple drug use -> '{display}'")
            else:
                # No substance data available, use generic
                display = find_icd10_entry("Mental and behavioural disorders due to multiple drug use")
                if display and display not in extracted_diagnoses:
                    extracted_diagnoses.append(display)
                    print(f"[GPR-DX] SUBSTANCE: No section 8 data, using multiple drug use -> '{display}'")

        # === STEP 3: Only use notes if we still need more diagnoses ===
        # Notes search is AFTER substance detection to ensure substance gets priority
        if len(extracted_diagnoses) < 3 and raw_notes:
            print(f"[GPR-DX] Checking {len(raw_notes)} notes for additional diagnoses")
            all_text = ""
            for note in raw_notes:
                content = note.get("content", "") or note.get("text", "") or ""
                all_text += " " + content

            if all_text.strip():
                extract_from_text(all_text, "NOTES")

        print(f"[GPR-DX] Total extracted: {len(extracted_diagnoses)} diagnoses")

        if not extracted_diagnoses:
            print(f"[GPR-DX] No diagnoses found")
            return

        # Populate the comboboxes
        for i, dx_name in enumerate(extracted_diagnoses):
            if i >= len(self.dx_boxes):
                break

            combo = self.dx_boxes[i]
            index = combo.findText(dx_name)
            if index >= 0:
                combo.blockSignals(True)
                combo.setCurrentIndex(index)
                combo.blockSignals(False)
                print(f"[GPR-DX] Set diagnosis {i+1}: {dx_name}")

        self._update_preview()
        print(f"[GPR-DX] ✓ Pre-filled {len(extracted_diagnoses)} diagnosis(es)")

    def get_state(self) -> dict:
        """Get current state for saving."""
        dx_list = []
        for combo in self.dx_boxes:
            meta = combo.currentData()
            if meta:
                dx_list.append(meta)
        return {"diagnoses": dx_list}

    def load_state(self, state: dict):
        """Load state from saved data."""
        for combo in self.dx_boxes:
            combo.blockSignals(True)
            combo.setCurrentIndex(0)
            combo.blockSignals(False)

        for combo, meta in zip(self.dx_boxes, state.get("diagnoses", [])):
            if not meta:
                continue
            index = combo.findText(meta.get("diagnosis", ""))
            if index >= 0:
                combo.blockSignals(True)
                combo.setCurrentIndex(index)
                combo.blockSignals(False)

        self._update_preview()


# ================================================================
# GPR LEGAL CRITERIA POPUP (Section 12)
# ================================================================

class GPRLegalCriteriaPopup(QWidget):
    """Legal Criteria popup with mental disorder, nature/degree, necessity, and treatment sections."""

    sent = Signal(str)

    def __init__(self, parent=None, gender=None, icd10_dict=None):
        super().__init__(parent)
        self.gender = gender or "neutral"
        self.icd10_dict = icd10_dict or {}
        print(f"[GPRLegalCriteriaPopup] __init__: gender={self.gender}, icd10_dict has {len(self.icd10_dict)} entries")
        self._diagnosis_text = ""  # Will be set from section 11 or combo
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._extracted_checkboxes = []
        self._setup_ui()
        print(f"[GPRLegalCriteriaPopup] After _setup_ui: dx_combos={len(self.dx_combos)}, items_in_first_combo={self.dx_combos[0].count() if self.dx_combos else 'N/A'}")
        add_lock_to_popup(self, show_button=False)

    def eventFilter(self, obj, event):
        """Block wheel events on combo boxes to prevent accidental scrolling changes."""
        from PySide6.QtCore import QEvent
        from PySide6.QtWidgets import QComboBox
        if event.type() == QEvent.Type.Wheel and isinstance(obj, QComboBox):
            return True
        return super().eventFilter(obj, event)

    def _get_pronouns(self):
        g = (self.gender or "").lower().strip()
        if g == "male":
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "his", "himself": "himself", "suffers": "suffers", "does": "does"}
        elif g == "female":
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "her", "himself": "herself", "suffers": "suffers", "does": "does"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "their", "himself": "themselves", "suffers": "suffer", "does": "do"}

    def set_gender(self, gender: str):
        self.gender = gender
        self._update_preview()

    def update_gender(self, gender: str):
        """Alias for set_gender for consistency with other popups."""
        self.set_gender(gender)

    def set_diagnosis(self, diagnosis_text: str):
        """Set the diagnosis text from section 11."""
        self._diagnosis_text = diagnosis_text
        self._update_preview()

    def _setup_ui(self):
        from PySide6.QtWidgets import QRadioButton, QButtonGroup, QSlider, QTextEdit

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        self.main_layout = QVBoxLayout(main_container)
        self.main_layout.setContentsMargins(4, 4, 4, 4)
        self.main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: INPUT FORM (collapsible with drag bar)
        # ====================================================
        self.input_section = CollapsibleSection("Legal Criteria", start_collapsed=False)
        self.input_section.set_content_height(800)
        self.input_section._min_height = 200
        self.input_section._max_height = 1200
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.input_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #005a50;
                background: transparent;
                border: none;
            }
        """)

        # Form scroll area
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255,255,255,0.95);
                border: 1px solid rgba(0, 140, 126, 0.2);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        form_container = QWidget()
        form_container.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(12, 12, 12, 12)
        form_layout.setSpacing(16)

        # ============================================
        # 1. MENTAL DISORDER - Present/Absent
        # ============================================
        md_label = QLabel("Mental Disorder")
        md_label.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent;")
        form_layout.addWidget(md_label)

        self.md_group = QButtonGroup(self)
        md_row = QHBoxLayout()
        md_row.setSpacing(16)

        self.md_present = QRadioButton("Present")
        self.md_present.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.md_absent = QRadioButton("Absent")
        self.md_absent.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.md_group.addButton(self.md_present, 0)
        self.md_group.addButton(self.md_absent, 1)

        self.md_present.toggled.connect(self._on_md_toggled)
        self.md_absent.toggled.connect(self._update_preview)

        md_row.addWidget(self.md_present)
        md_row.addWidget(self.md_absent)
        md_row.addStretch()
        form_layout.addLayout(md_row)

        # ICD-10 diagnosis combo (shown when Present is selected)
        self.dx_container = QWidget()
        self.dx_container.setStyleSheet("background: transparent;")
        dx_layout = QVBoxLayout(self.dx_container)
        dx_layout.setContentsMargins(0, 4, 0, 0)
        dx_layout.setSpacing(6)

        dx_lbl = QLabel("Diagnosis (ICD-10)")
        dx_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151; background: transparent;")
        dx_layout.addWidget(dx_lbl)

        from PySide6.QtWidgets import QComboBox, QCompleter, QStyleFactory
        self.dx_combos = []
        for i in range(3):
            combo = QComboBox()
            combo.setStyle(QStyleFactory.create("Fusion"))
            combo.setEditable(True)
            combo.lineEdit().setReadOnly(True)
            combo.lineEdit().setPlaceholderText("Start typing to search...")
            combo.setSizeAdjustPolicy(QComboBox.SizeAdjustPolicy.AdjustToMinimumContentsLengthWithIcon)
            combo.setMinimumContentsLength(25)
            combo.setStyleSheet("""
                QComboBox {
                    padding: 6px;
                    font-size: 17px;
                    background: white;
                    border: 1px solid #d1d5db;
                    border-radius: 6px;
                }
                QComboBox QAbstractItemView {
                    min-width: 400px;
                }
            """)
            combo.addItem("Not specified", None)
            for diagnosis, meta in sorted(
                self.icd10_dict.items(),
                key=lambda x: x[0].lower()
            ):
                icd_code = meta.get("icd10") if isinstance(meta, dict) else meta
                combo.addItem(diagnosis, {"diagnosis": diagnosis, "icd10": icd_code})

            completer = QCompleter(combo.model(), combo)
            completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
            completer.setFilterMode(Qt.MatchFlag.MatchContains)
            completer.setCompletionMode(QCompleter.CompletionMode.PopupCompletion)
            combo.setCompleter(completer)
            combo.setMaxVisibleItems(15)
            combo.currentIndexChanged.connect(self._update_preview)
            combo.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
            combo.installEventFilter(self)

            dx_layout.addWidget(combo)
            self.dx_combos.append(combo)

        self.dx_container.hide()
        form_layout.addWidget(self.dx_container)

        # Container for criteria (shown when Present is selected)
        self.criteria_container = QWidget()
        self.criteria_container.setStyleSheet("background: transparent;")
        criteria_layout = QVBoxLayout(self.criteria_container)
        criteria_layout.setContentsMargins(0, 8, 0, 0)
        criteria_layout.setSpacing(12)

        # ============================================
        # 2. CRITERIA WARRANTING DETENTION - Met/Not Met
        # ============================================
        cwd_label = QLabel("Criteria Warranting Detention")
        cwd_label.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent;")
        criteria_layout.addWidget(cwd_label)

        self.cwd_group = QButtonGroup(self)
        cwd_row = QHBoxLayout()
        cwd_row.setSpacing(16)

        self.cwd_met = QRadioButton("Met")
        self.cwd_met.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.cwd_not_met = QRadioButton("Not Met")
        self.cwd_not_met.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.cwd_group.addButton(self.cwd_met, 0)
        self.cwd_group.addButton(self.cwd_not_met, 1)

        self.cwd_met.toggled.connect(self._on_cwd_toggled)
        self.cwd_not_met.toggled.connect(self._update_preview)

        cwd_row.addWidget(self.cwd_met)
        cwd_row.addWidget(self.cwd_not_met)
        cwd_row.addStretch()
        criteria_layout.addLayout(cwd_row)

        # Nature/Degree container (shown when Met is selected)
        self.nature_degree_container = QWidget()
        self.nature_degree_container.setStyleSheet("background: transparent;")
        nd_layout = QVBoxLayout(self.nature_degree_container)
        nd_layout.setContentsMargins(16, 8, 0, 0)
        nd_layout.setSpacing(8)

        # Nature checkbox
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151; background: transparent;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        nd_layout.addWidget(self.nature_cb)

        # Nature sub-options container
        self.nature_options_container = QWidget()
        self.nature_options_container.setStyleSheet("background: transparent;")
        nature_opt_layout = QVBoxLayout(self.nature_options_container)
        nature_opt_layout.setContentsMargins(24, 4, 0, 0)
        nature_opt_layout.setSpacing(4)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.relapsing_cb.toggled.connect(self._update_preview)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.treatment_resistant_cb.toggled.connect(self._update_preview)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.chronic_cb.toggled.connect(self._update_preview)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options_container.hide()
        nd_layout.addWidget(self.nature_options_container)

        # Degree checkbox
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151; background: transparent;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        nd_layout.addWidget(self.degree_cb)

        # Degree sub-options container
        self.degree_options_container = QWidget()
        self.degree_options_container.setStyleSheet("background: transparent;")
        degree_opt_layout = QVBoxLayout(self.degree_options_container)
        degree_opt_layout.setContentsMargins(24, 4, 0, 0)
        degree_opt_layout.setSpacing(8)

        # Severity slider
        slider_label = QLabel("Symptom severity:")
        slider_label.setStyleSheet("font-size: 16px; color: #6b7280; background: transparent;")
        degree_opt_layout.addWidget(slider_label)

        slider_row = QHBoxLayout()
        slider_row.setSpacing(8)

        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(200)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500; margin-left: 0px; background: transparent;")
        degree_opt_layout.addWidget(self.degree_level_label)

        # Details text box
        details_label = QLabel("Symptoms including:")
        details_label.setStyleSheet("font-size: 16px; color: #6b7280; background: transparent;")
        degree_opt_layout.addWidget(details_label)

        self.degree_details = QTextEdit()
        self.degree_details.setPlaceholderText("Enter symptom details...")
        self.degree_details.setMaximumHeight(80)
        self.degree_details.setStyleSheet("""
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 6px;
                background: white;
                font-size: 18px;
            }
        """)
        self.degree_details.textChanged.connect(self._update_preview)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options_container.hide()
        nd_layout.addWidget(self.degree_options_container)

        self.nature_degree_container.hide()
        criteria_layout.addWidget(self.nature_degree_container)

        # ============================================
        # 3. NECESSITY SECTION (from Section 20)
        # ============================================
        necessity_label = QLabel("Necessity")
        necessity_label.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent; margin-top: 8px;")
        criteria_layout.addWidget(necessity_label)

        # Necessary Yes/No
        self.necessary_group = QButtonGroup(self)
        necessary_row = QHBoxLayout()
        necessary_row.setSpacing(16)

        self.necessary_yes = QRadioButton("Yes")
        self.necessary_yes.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.necessary_no = QRadioButton("No")
        self.necessary_no.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.necessary_group.addButton(self.necessary_yes, 0)
        self.necessary_group.addButton(self.necessary_no, 1)

        self.necessary_yes.toggled.connect(self._on_necessary_toggled)
        self.necessary_no.toggled.connect(self._update_preview)

        necessary_row.addWidget(self.necessary_yes)
        necessary_row.addWidget(self.necessary_no)
        necessary_row.addStretch()
        criteria_layout.addLayout(necessary_row)

        # Health & Safety container
        self.health_safety_container = QWidget()
        self.health_safety_container.setStyleSheet("background: transparent;")
        hs_layout = QVBoxLayout(self.health_safety_container)
        hs_layout.setContentsMargins(16, 8, 0, 0)
        hs_layout.setSpacing(8)

        # Health checkbox
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151; background: transparent;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        hs_layout.addWidget(self.health_cb)

        # Health sub-options
        self.health_container = QWidget()
        self.health_container.setStyleSheet("background: transparent;")
        health_layout = QVBoxLayout(self.health_container)
        health_layout.setContentsMargins(24, 4, 0, 0)
        health_layout.setSpacing(4)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_layout.addWidget(self.mental_health_cb)

        # Mental Health sub-options
        self.mental_health_container = QWidget()
        self.mental_health_container.setStyleSheet("background: transparent;")
        mh_layout = QVBoxLayout(self.mental_health_container)
        mh_layout.setContentsMargins(24, 4, 0, 0)
        mh_layout.setSpacing(4)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.poor_compliance_cb.toggled.connect(self._update_preview)
        mh_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.limited_insight_cb.toggled.connect(self._update_preview)
        mh_layout.addWidget(self.limited_insight_cb)

        self.mental_health_container.hide()
        health_layout.addWidget(self.mental_health_container)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QTextEdit()
        self.physical_health_details.setPlaceholderText("Enter physical health details...")
        self.physical_health_details.setMaximumHeight(60)
        self.physical_health_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 4px; background: white; font-size: 18px;")
        self.physical_health_details.textChanged.connect(self._update_preview)
        self.physical_health_details.hide()
        health_layout.addWidget(self.physical_health_details)

        self.health_container.hide()
        hs_layout.addWidget(self.health_container)

        # Safety checkbox
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151; background: transparent;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        hs_layout.addWidget(self.safety_cb)

        # Safety sub-options
        self.safety_container = QWidget()
        self.safety_container.setStyleSheet("background: transparent;")
        safety_layout = QVBoxLayout(self.safety_container)
        safety_layout.setContentsMargins(24, 4, 0, 0)
        safety_layout.setSpacing(4)

        self.self_cb = QCheckBox("Self")
        self.self_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.self_cb.toggled.connect(self._on_self_toggled)
        safety_layout.addWidget(self.self_cb)

        self.self_details = QTextEdit()
        self.self_details.setPlaceholderText("Enter details about risk to self...")
        self.self_details.setMaximumHeight(60)
        self.self_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 4px; background: white; font-size: 18px;")
        self.self_details.textChanged.connect(self._update_preview)
        self.self_details.hide()
        safety_layout.addWidget(self.self_details)

        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("""
            QRadioButton {
                font-size: 18px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_layout.addWidget(self.others_cb)

        self.others_details = QTextEdit()
        self.others_details.setPlaceholderText("Enter details about risk to others...")
        self.others_details.setMaximumHeight(60)
        self.others_details.setStyleSheet("border: 1px solid #d1d5db; border-radius: 4px; padding: 4px; background: white; font-size: 18px;")
        self.others_details.textChanged.connect(self._update_preview)
        self.others_details.hide()
        safety_layout.addWidget(self.others_details)

        self.safety_container.hide()
        hs_layout.addWidget(self.safety_container)

        self.health_safety_container.hide()
        criteria_layout.addWidget(self.health_safety_container)

        # ============================================
        # 4. TREATMENT AVAILABLE
        # ============================================
        self.treatment_cb = QCheckBox("Treatment Available")
        self.treatment_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent; margin-top: 8px;")
        self.treatment_cb.toggled.connect(self._update_preview)
        criteria_layout.addWidget(self.treatment_cb)

        # ============================================
        # 5. LEAST RESTRICTIVE
        # ============================================
        self.least_restrictive_cb = QCheckBox("Least Restrictive Option")
        self.least_restrictive_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent;")
        self.least_restrictive_cb.toggled.connect(self._update_preview)
        criteria_layout.addWidget(self.least_restrictive_cb)

        self.criteria_container.hide()
        form_layout.addWidget(self.criteria_container)

        form_layout.addStretch()
        form_scroll.setWidget(form_container)
        self.input_section.set_content(form_scroll)
        self.main_layout.addWidget(self.input_section)

        # ====================================================
        # SECTION 3: EXTRACTED DATA (collapsible, at bottom)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 400
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        self.extracted_scroll = QScrollArea()
        self.extracted_scroll.setWidgetResizable(True)
        self.extracted_scroll.setFrameShape(QScrollArea.NoFrame)
        self.extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.extracted_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(12, 12, 12, 12)
        self.extracted_checkboxes_layout.setSpacing(12)

        self.extracted_scroll.setWidget(self.extracted_container)
        self.extracted_section.set_content(self.extracted_scroll)
        self.extracted_section.setVisible(False)
        self.main_layout.addWidget(self.extracted_section)

        self.main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    # ============================================
    # TOGGLE HANDLERS
    # ============================================
    def _on_md_toggled(self, checked):
        try:
            present = self.md_present.isChecked()
            print(f"[GPRLegalCriteriaPopup] _on_md_toggled: checked={checked}, present={present}, icd10_items={len(self.icd10_dict)}")
            self.dx_container.setVisible(present)
            self.criteria_container.setVisible(present)
            # Force layout recalculation (nested scroll areas can fail to update on Windows)
            if present:
                self.dx_container.updateGeometry()
                self.criteria_container.updateGeometry()
                parent = self.dx_container.parentWidget()
                while parent and parent is not self:
                    if parent.layout():
                        parent.layout().activate()
                    parent.updateGeometry()
                    parent = parent.parentWidget()
                from PySide6.QtWidgets import QApplication
                QApplication.processEvents()
                # Scroll dx_container into view
                from PySide6.QtCore import QTimer
                QTimer.singleShot(100, lambda: self._scroll_to_widget(self.dx_container))
            print(f"[GPRLegalCriteriaPopup] dx_container visible={self.dx_container.isVisible()}, size={self.dx_container.size().width()}x{self.dx_container.size().height()}")
        except Exception as e:
            print(f"[GPRLegalCriteriaPopup] ERROR in _on_md_toggled: {e}")
            import traceback
            traceback.print_exc()
        self._update_preview()

    def _scroll_to_widget(self, widget):
        """Scroll the nearest parent QScrollArea to make widget visible."""
        from PySide6.QtWidgets import QScrollArea
        parent = widget.parentWidget()
        while parent:
            if isinstance(parent, QScrollArea):
                parent.ensureWidgetVisible(widget, 50, 50)
                break
            parent = parent.parentWidget()

    def _on_cwd_toggled(self, checked):
        if self.cwd_met.isChecked():
            self.nature_degree_container.show()
        else:
            self.nature_degree_container.hide()
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
        self._update_preview()

    def _on_nature_toggled(self, checked):
        self.nature_options_container.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_preview()

    def _on_degree_toggled(self, checked):
        self.degree_options_container.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_preview()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_preview()

    def _on_necessary_toggled(self, checked):
        if self.necessary_yes.isChecked():
            self.health_safety_container.show()
        else:
            self.health_safety_container.hide()
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
        self._update_preview()

    def _on_health_toggled(self, checked):
        self.health_container.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_preview()

    def _on_mental_health_toggled(self, checked):
        self.mental_health_container.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_preview()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_preview()

    def _on_safety_toggled(self, checked):
        self.safety_container.setVisible(checked)
        if not checked:
            self.self_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_preview()

    def _on_self_toggled(self, checked):
        self.self_details.setVisible(checked)
        if not checked:
            self.self_details.clear()
        self._update_preview()

    def _on_others_toggled(self, checked):
        self.others_details.setVisible(checked)
        if not checked:
            self.others_details.clear()
        self._update_preview()

    # ============================================
    # TEXT GENERATION
    # ============================================
    def _update_preview(self):
        """Build text from selections and send directly to card."""
        text = self.generate_text()
        # Send directly to card
        self.sent.emit(text)

    def _get_selected_diagnoses(self) -> list:
        """Get list of selected diagnosis strings from ICD-10 combos."""
        diagnoses = []
        for combo in self.dx_combos:
            meta = combo.currentData()
            if meta and isinstance(meta, dict):
                dx = meta.get("diagnosis", "")
                icd = meta.get("icd10", "")
                if dx:
                    diagnoses.append(f"{dx} ({icd})" if icd else dx)
        return diagnoses

    def generate_text(self) -> str:
        p = self._get_pronouns()
        parts = []

        # 1. Mental Disorder + Nature/Degree (combined into one sentence)
        if self.md_present.isChecked():
            # Build diagnosis text from combo selection, then fallback to external set_diagnosis
            combo_diagnoses = self._get_selected_diagnoses()
            if combo_diagnoses:
                dx_text = ", ".join(combo_diagnoses[:-1]) + " and " + combo_diagnoses[-1] if len(combo_diagnoses) > 1 else combo_diagnoses[0]
            elif self._diagnosis_text:
                dx_text = self._diagnosis_text
            else:
                dx_text = ""

            if dx_text:
                md_base = f"{p['subj']} {p['suffers']} from {dx_text} which is a mental disorder under the Mental Health Act"
            else:
                md_base = f"{p['subj']} {p['suffers']} from a mental disorder under the Mental Health Act"

            # Check for nature/degree to append
            if self.cwd_met.isChecked():
                nature_checked = self.nature_cb.isChecked()
                degree_checked = self.degree_cb.isChecked()

                if nature_checked and degree_checked:
                    nd_text = ", which is of a nature and degree to warrant detention."
                elif nature_checked:
                    nd_text = ", which is of a nature to warrant detention."
                elif degree_checked:
                    nd_text = ", which is of a degree to warrant detention."
                else:
                    nd_text = "."

                parts.append(md_base + nd_text)

                # Nature sub-options
                if nature_checked:
                    nature_types = []
                    if self.relapsing_cb.isChecked():
                        nature_types.append("relapsing and remitting")
                    if self.treatment_resistant_cb.isChecked():
                        nature_types.append("treatment resistant")
                    if self.chronic_cb.isChecked():
                        nature_types.append("chronic and enduring")

                    if nature_types:
                        nature_str = ", ".join(nature_types)
                        parts.append(f"The illness is of a {nature_str} nature.")

                # Degree sub-options
                if degree_checked:
                    levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                    level = levels.get(self.degree_slider.value(), "several")
                    details = self.degree_details.toPlainText().strip()
                    if details:
                        parts.append(f"The degree of the illness is evidenced by {level} symptoms including {details}.")
                    else:
                        parts.append(f"The degree of the illness is evidenced by {level} symptoms.")

            elif self.cwd_not_met.isChecked():
                parts.append(md_base + ".")
                parts.append("The criteria for detention are not met.")
            else:
                parts.append(md_base + ".")

        elif self.md_absent.isChecked():
            parts.append(f"{p['subj']} {p['does']} not suffer from a mental disorder under the Mental Health Act.")

        # 2. Necessity (combined "is necessary to prevent deterioration")
        if self.necessary_yes.isChecked():
            # Health - Mental Health
            if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
                parts.append(f"Medical treatment under the Mental Health Act is necessary to prevent deterioration in {p['pos']} mental health.")

                poor = self.poor_compliance_cb.isChecked()
                limited = self.limited_insight_cb.isChecked()

                if poor and limited:
                    parts.append(f"Both historical non compliance and current limited insight makes the risk on stopping medication high without the safeguards of the Mental Health Act. This would result in a deterioration of {p['pos']} mental state.")
                elif poor:
                    parts.append(f"This is based on historical non compliance and without detention I would be concerned this would result in a deterioration of {p['pos']} mental state.")
                elif limited:
                    parts.append(f"I am concerned about {p['pos']} current limited insight into {p['pos']} mental health needs and how this would result in immediate non compliance with medication, hence a deterioration in {p['pos']} mental health.")
            else:
                parts.append("Medical treatment under the Mental Health Act is necessary.")

            # Health - Physical Health
            if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
                details = self.physical_health_details.toPlainText().strip()
                if self.mental_health_cb.isChecked():
                    base = f"The Mental Health Act is also necessary for maintaining {p['pos']} physical health."
                else:
                    base = f"The Mental Health Act is necessary for {p['pos']} physical health."
                if details:
                    parts.append(f"{base} {details}")
                else:
                    parts.append(base)

            # Safety
            if self.safety_cb.isChecked():
                # Determine if we need "also" - use it if health is checked
                use_also = self.health_cb.isChecked()

                if self.self_cb.isChecked():
                    details = self.self_details.toPlainText().strip()
                    if use_also:
                        base = f"The Mental Health Act is also necessary for {p['pos']} risk to {p['himself']}."
                    else:
                        base = f"The Mental Health Act is necessary for {p['pos']} risk to {p['himself']}."
                    if details:
                        parts.append(f"{base} {details}")
                    else:
                        parts.append(base)

                if self.others_cb.isChecked():
                    details = self.others_details.toPlainText().strip()
                    # Use "also" if health is checked OR if self is checked
                    if use_also or self.self_cb.isChecked():
                        base = "Risk to others also makes the Mental Health Act necessary."
                    else:
                        base = "Risk to others makes the Mental Health Act necessary."
                    if details:
                        parts.append(f"{base} {details}")
                    else:
                        parts.append(base)

        elif self.necessary_no.isChecked():
            parts.append("Medical treatment under the Mental Health Act is not necessary.")

        # 3. Treatment Available
        if self.treatment_cb.isChecked():
            parts.append("Treatment is available, medical, nursing, OT/Psychology and social work.")

        # 4. Least Restrictive
        if self.least_restrictive_cb.isChecked():
            parts.append(f"I can confirm this is the least restrictive option to meet {p['pos']} needs.")

        # 5. Include checked extracted entries
        extracted_parts = []
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                full_text = cb.property("full_text")
                if full_text:
                    extracted_parts.append(full_text)

        if extracted_parts:
            parts.append("\n\n" + "\n\n".join(extracted_parts))

        return " ".join(parts)

    def _send_to_report(self):
        text = self.preview_label.text()
        if text and text != "Select options...":
            self.sent.emit(text)

    def set_entries(self, entries: list):
        """Set entries from extracted data with checkboxes."""
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        has_entries = False
        for entry in entries:
            text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
            if text and text.strip():
                has_entries = True

                entry_widget = QWidget()
                entry_widget.setStyleSheet("background: transparent;")
                entry_layout = QHBoxLayout(entry_widget)
                entry_layout.setContentsMargins(0, 4, 0, 4)
                entry_layout.setSpacing(8)

                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(20, 20)
                cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.stateChanged.connect(self._update_preview)
                entry_layout.addWidget(cb, 0, Qt.AlignmentFlag.AlignTop)

                label = QLabel(text)
                label.setWordWrap(True)
                label.setStyleSheet("QLabel { background: transparent; border: none; font-size: 17px; color: #4a4a4a; }")
                entry_layout.addWidget(label, 1)

                self.extracted_checkboxes_layout.addWidget(entry_widget)
                self._extracted_checkboxes.append(cb)

        if has_entries:
            self.extracted_section.setVisible(True)
            if self.extracted_section.is_collapsed():
                self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)

    def get_state(self) -> dict:
        dx_list = []
        for combo in self.dx_combos:
            meta = combo.currentData()
            if meta:
                dx_list.append(meta)
        return {
            "md_present": self.md_present.isChecked(),
            "md_absent": self.md_absent.isChecked(),
            "diagnoses": dx_list,
            "cwd_met": self.cwd_met.isChecked(),
            "cwd_not_met": self.cwd_not_met.isChecked(),
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.toPlainText(),
            "necessary_yes": self.necessary_yes.isChecked(),
            "necessary_no": self.necessary_no.isChecked(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.toPlainText(),
            "safety": self.safety_cb.isChecked(),
            "self": self.self_cb.isChecked(),
            "self_details": self.self_details.toPlainText(),
            "others": self.others_cb.isChecked(),
            "others_details": self.others_details.toPlainText(),
            "treatment": self.treatment_cb.isChecked(),
            "least_restrictive": self.least_restrictive_cb.isChecked(),
        }

    def load_state(self, state: dict):
        if not state:
            return

        if state.get("md_present"):
            self.md_present.setChecked(True)
        elif state.get("md_absent"):
            self.md_absent.setChecked(True)

        # Restore ICD-10 combo selections
        dx_list = state.get("diagnoses", [])
        for i, combo in enumerate(self.dx_combos):
            if i < len(dx_list) and dx_list[i]:
                meta = dx_list[i]
                dx_name = meta.get("diagnosis", "")
                for j in range(combo.count()):
                    d = combo.itemData(j)
                    if d and isinstance(d, dict) and d.get("diagnosis") == dx_name:
                        combo.setCurrentIndex(j)
                        break

        if state.get("cwd_met"):
            self.cwd_met.setChecked(True)
        elif state.get("cwd_not_met"):
            self.cwd_not_met.setChecked(True)

        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))

        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setPlainText(state.get("degree_details", ""))

        if state.get("necessary_yes"):
            self.necessary_yes.setChecked(True)
        elif state.get("necessary_no"):
            self.necessary_no.setChecked(True)

        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setPlainText(state.get("physical_health_details", ""))

        self.safety_cb.setChecked(state.get("safety", False))
        self.self_cb.setChecked(state.get("self", False))
        self.self_details.setPlainText(state.get("self_details", ""))
        self.others_cb.setChecked(state.get("others", False))
        self.others_details.setPlainText(state.get("others_details", ""))

        self.treatment_cb.setChecked(state.get("treatment", False))
        self.least_restrictive_cb.setChecked(state.get("least_restrictive", False))

        self._update_preview()


# ================================================================
# GPR STRENGTHS POPUP (Section 13)
# ================================================================

class GPRStrengthsPopup(QWidget):
    """Strengths popup with preview, input checkboxes, and extracted data sections."""

    sent = Signal(str)

    def __init__(self, parent=None, gender=None):
        super().__init__(parent)
        self.gender = gender or "neutral"
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._extracted_checkboxes = []
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _get_pronouns(self):
        g = (self.gender or "").lower().strip()
        if g == "male":
            return {"subj": "He", "obj": "him", "pos": "His", "pos_lower": "his"}
        elif g == "female":
            return {"subj": "She", "obj": "her", "pos": "Her", "pos_lower": "her"}
        return {"subj": "They", "obj": "them", "pos": "Their", "pos_lower": "their"}

    def set_gender(self, gender: str):
        self.gender = gender
        self._update_preview()

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: STRENGTHS INPUT (collapsible with drag bar)
        # ====================================================
        self.input_section = CollapsibleSection("Strengths or Positive Factors", start_collapsed=True)
        self.input_section.set_content_height(350)
        self.input_section._min_height = 150
        self.input_section._max_height = 500
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.input_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #005a50;
                background: transparent;
                border: none;
            }
        """)

        # Form scroll area
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255,255,255,0.95);
                border: 1px solid rgba(0, 140, 126, 0.2);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        form_container = QWidget()
        form_container.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(12, 12, 12, 12)
        form_layout.setSpacing(8)

        # === ENGAGEMENT SECTION ===
        engagement_lbl = QLabel("Engagement")
        engagement_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent;")
        form_layout.addWidget(engagement_lbl)

        self.staff_cb = QCheckBox("Staff")
        self.staff_cb.toggled.connect(self._update_preview)
        form_layout.addWidget(self.staff_cb)

        self.peers_cb = QCheckBox("Peers")
        self.peers_cb.toggled.connect(self._update_preview)
        form_layout.addWidget(self.peers_cb)

        # === ACTIVITIES SECTION ===
        activities_lbl = QLabel("Activities & Treatment")
        activities_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent; margin-top: 8px;")
        form_layout.addWidget(activities_lbl)

        self.ot_cb = QCheckBox("OT")
        self.ot_cb.toggled.connect(self._update_preview)
        form_layout.addWidget(self.ot_cb)

        self.nursing_cb = QCheckBox("Nursing")
        self.nursing_cb.toggled.connect(self._update_preview)
        form_layout.addWidget(self.nursing_cb)

        self.psychology_cb = QCheckBox("Psychology")
        self.psychology_cb.toggled.connect(self._update_preview)
        form_layout.addWidget(self.psychology_cb)

        # === AFFECT SECTION ===
        affect_lbl = QLabel("Affect")
        affect_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; background: transparent; margin-top: 8px;")
        form_layout.addWidget(affect_lbl)

        self.affect_cb = QCheckBox("Affect (expand for options)")
        self.affect_cb.toggled.connect(self._on_affect_toggled)
        form_layout.addWidget(self.affect_cb)

        # Affect sub-options container
        self.affect_container = QWidget()
        self.affect_container.setStyleSheet("background: transparent;")
        affect_sub_layout = QVBoxLayout(self.affect_container)
        affect_sub_layout.setContentsMargins(20, 4, 0, 0)
        affect_sub_layout.setSpacing(4)

        self.humour_cb = QCheckBox("Sense of humour")
        self.humour_cb.toggled.connect(self._update_preview)
        affect_sub_layout.addWidget(self.humour_cb)

        self.warmth_cb = QCheckBox("Warmth")
        self.warmth_cb.toggled.connect(self._update_preview)
        affect_sub_layout.addWidget(self.warmth_cb)

        self.friendly_cb = QCheckBox("Friendly")
        self.friendly_cb.toggled.connect(self._update_preview)
        affect_sub_layout.addWidget(self.friendly_cb)

        self.caring_cb = QCheckBox("Caring")
        self.caring_cb.toggled.connect(self._update_preview)
        affect_sub_layout.addWidget(self.caring_cb)

        self.affect_container.hide()
        form_layout.addWidget(self.affect_container)

        form_layout.addStretch()
        form_scroll.setWidget(form_container)
        self.input_section.set_content(form_scroll)
        main_layout.addWidget(self.input_section)

        # ====================================================
        # SECTION 3: EXTRACTED DATA (collapsible, at bottom)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 400
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        self.extracted_scroll = QScrollArea()
        self.extracted_scroll.setWidgetResizable(True)
        self.extracted_scroll.setFrameShape(QScrollArea.NoFrame)
        self.extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.extracted_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(12, 12, 12, 12)
        self.extracted_checkboxes_layout.setSpacing(12)

        self.extracted_scroll.setWidget(self.extracted_container)
        self.extracted_section.set_content(self.extracted_scroll)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _on_affect_toggled(self, checked):
        self.affect_container.setVisible(checked)
        if not checked:
            self.humour_cb.setChecked(False)
            self.warmth_cb.setChecked(False)
            self.friendly_cb.setChecked(False)
            self.caring_cb.setChecked(False)
        self._update_preview()

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        text = self.generate_text()
        # Add extracted checked items
        extracted_parts = []
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                extracted_parts.append(cb.property("full_text"))

        if text and extracted_parts:
            combined = text + " " + " ".join(extracted_parts)
        elif text:
            combined = text
        elif extracted_parts:
            combined = " ".join(extracted_parts)
        else:
            combined = ""

        # Send directly to card
        self.sent.emit(combined)

    def generate_text(self) -> str:
        p = self._get_pronouns()
        parts = []

        # Engagement with staff/peers
        staff = self.staff_cb.isChecked()
        peers = self.peers_cb.isChecked()

        if staff and peers:
            parts.append(f"{p['subj']} engages well with both staff and peers.")
        elif staff:
            parts.append(f"{p['subj']} engages well with staff.")
        elif peers:
            parts.append(f"{p['subj']} engages well with peers.")

        # OT
        if self.ot_cb.isChecked():
            parts.append(f"{p['subj']} is able to get involved in OT activities.")

        # Nursing
        if self.nursing_cb.isChecked():
            parts.append(f"{p['subj']} works collaboratively with nursing staff.")

        # Psychology
        if self.psychology_cb.isChecked():
            parts.append(f"{p['pos']} attendance at psychology sessions is an important strength.")

        # Affect sub-options
        if self.affect_cb.isChecked():
            if self.humour_cb.isChecked():
                parts.append(f"{p['subj']} can display a positive sense of humour.")
            if self.warmth_cb.isChecked():
                parts.append(f"{p['subj']} can be warm with staff and peers.")
            if self.friendly_cb.isChecked():
                parts.append(f"{p['subj']} can be appropriately friendly at times.")
            if self.caring_cb.isChecked():
                parts.append(f"{p['subj']} displays empathy and a caring attitude on the ward to staff and peers.")

        return " ".join(parts)

    def set_entries(self, entries: list):
        """Set entries from extracted data with checkboxes."""
        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        has_entries = False
        for entry in entries:
            text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
            if text and text.strip():
                has_entries = True

                entry_widget = QWidget()
                entry_widget.setStyleSheet("background: transparent;")
                entry_layout = QHBoxLayout(entry_widget)
                entry_layout.setContentsMargins(0, 4, 0, 4)
                entry_layout.setSpacing(8)

                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(20, 20)
                cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.stateChanged.connect(self._update_preview)
                entry_layout.addWidget(cb, 0, Qt.AlignmentFlag.AlignTop)

                label = QLabel(text)
                label.setWordWrap(True)
                label.setStyleSheet("QLabel { background: transparent; border: none; font-size: 17px; color: #4a4a4a; }")
                entry_layout.addWidget(label, 1)

                self.extracted_checkboxes_layout.addWidget(entry_widget)
                self._extracted_checkboxes.append(cb)

        if has_entries:
            self.extracted_section.setVisible(True)
            if self.extracted_section.is_collapsed():
                self.extracted_section._toggle_collapse()
        else:
            self.extracted_section.setVisible(False)

    def get_state(self) -> dict:
        return {
            "staff": self.staff_cb.isChecked(),
            "peers": self.peers_cb.isChecked(),
            "ot": self.ot_cb.isChecked(),
            "nursing": self.nursing_cb.isChecked(),
            "psychology": self.psychology_cb.isChecked(),
            "affect": self.affect_cb.isChecked(),
            "humour": self.humour_cb.isChecked(),
            "warmth": self.warmth_cb.isChecked(),
            "friendly": self.friendly_cb.isChecked(),
            "caring": self.caring_cb.isChecked(),
        }

    def load_state(self, state: dict):
        if not state:
            return

        self.staff_cb.setChecked(state.get("staff", False))
        self.peers_cb.setChecked(state.get("peers", False))
        self.ot_cb.setChecked(state.get("ot", False))
        self.nursing_cb.setChecked(state.get("nursing", False))
        self.psychology_cb.setChecked(state.get("psychology", False))
        self.affect_cb.setChecked(state.get("affect", False))
        self.humour_cb.setChecked(state.get("humour", False))
        self.warmth_cb.setChecked(state.get("warmth", False))
        self.friendly_cb.setChecked(state.get("friendly", False))
        self.caring_cb.setChecked(state.get("caring", False))

        self._update_preview()


# ================================================================
# GPR SIGNATURE POPUP
# ================================================================

class GPRSignaturePopup(QWidget):
    """Signature popup with preview, input fields, and auto-load from mydetails."""

    sent = Signal(str)

    def __init__(self, parent=None, my_details=None):
        super().__init__(parent)
        self.my_details = my_details or {}
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._setup_ui()
        self._prefill_from_mydetails()
        add_lock_to_popup(self, show_button=False)

    def set_my_details(self, my_details: dict):
        """Set my_details and update fields."""
        self.my_details = my_details or {}
        self._prefill_from_mydetails()

    def _prefill_from_mydetails(self):
        """Pre-fill fields from MyDetails data."""
        if not self.my_details:
            return

        # Pre-fill name
        if self.my_details.get("full_name"):
            self.name_field.setText(self.my_details["full_name"])

        # Pre-fill designation (role_title)
        if self.my_details.get("role_title"):
            self.designation_field.setText(self.my_details["role_title"])

        # Pre-fill qualifications (discipline)
        if self.my_details.get("discipline"):
            self.qualifications_field.setText(self.my_details["discipline"])

        # Pre-fill GMC (registration_number)
        if self.my_details.get("registration_number"):
            self.gmc_field.setText(self.my_details["registration_number"])

        self._update_preview()

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: SIGNATURE DETAILS INPUT (collapsible)
        # ====================================================
        self.input_section = CollapsibleSection("Signature Details", start_collapsed=True)
        self.input_section.set_content_height(320)
        self.input_section._min_height = 150
        self.input_section._max_height = 450
        self.input_section.set_header_style("""
            QFrame {
                background: rgba(0, 140, 126, 0.15);
                border: 1px solid rgba(0, 140, 126, 0.3);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.input_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #005a50;
                background: transparent;
                border: none;
            }
        """)

        # Form scroll area
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("""
            QScrollArea {
                background: rgba(255,255,255,0.95);
                border: 1px solid rgba(0, 140, 126, 0.2);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
        """)

        form_container = QWidget()
        form_container.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(form_container)
        form_layout.setContentsMargins(12, 12, 12, 12)
        form_layout.setSpacing(10)

        # Common field style
        field_style = """
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 17px;
                color: #374151;
            }
            QLineEdit:focus {
                border-color: #008C7E;
            }
        """
        label_style = "font-size: 17px; font-weight: 600; color: #374151; background: transparent;"

        # Name
        name_lbl = QLabel("Name")
        name_lbl.setStyleSheet(label_style)
        form_layout.addWidget(name_lbl)

        self.name_field = QLineEdit()
        self.name_field.setPlaceholderText("Enter your full name")
        self.name_field.setStyleSheet(field_style)
        self.name_field.textChanged.connect(self._update_preview)
        form_layout.addWidget(self.name_field)

        # Designation
        designation_lbl = QLabel("Designation")
        designation_lbl.setStyleSheet(label_style)
        form_layout.addWidget(designation_lbl)

        self.designation_field = QLineEdit()
        self.designation_field.setPlaceholderText("e.g. Consultant Psychiatrist, Responsible Clinician")
        self.designation_field.setStyleSheet(field_style)
        self.designation_field.textChanged.connect(self._update_preview)
        form_layout.addWidget(self.designation_field)

        # Qualifications
        qualifications_lbl = QLabel("Qualifications")
        qualifications_lbl.setStyleSheet(label_style)
        form_layout.addWidget(qualifications_lbl)

        self.qualifications_field = QLineEdit()
        self.qualifications_field.setPlaceholderText("e.g. MBChB, MRCPsych, MD")
        self.qualifications_field.setStyleSheet(field_style)
        self.qualifications_field.textChanged.connect(self._update_preview)
        form_layout.addWidget(self.qualifications_field)

        # GMC/Professional Registration
        gmc_lbl = QLabel("GMC/Professional Registration Number")
        gmc_lbl.setStyleSheet(label_style)
        form_layout.addWidget(gmc_lbl)

        self.gmc_field = QLineEdit()
        self.gmc_field.setPlaceholderText("Enter registration number")
        self.gmc_field.setStyleSheet(field_style)
        self.gmc_field.textChanged.connect(self._update_preview)
        form_layout.addWidget(self.gmc_field)

        # Date
        date_lbl = QLabel("Date")
        date_lbl.setStyleSheet(label_style)
        form_layout.addWidget(date_lbl)

        self.date_field = QDateEdit()
        self.date_field.setDisplayFormat("dd/MM/yyyy")
        self.date_field.setCalendarPopup(True)
        self.date_field.setDate(QDate.currentDate())
        self.date_field.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 17px;
                color: #374151;
            }
            QDateEdit:focus {
                border-color: #008C7E;
            }
            QDateEdit::drop-down {
                border: none;
                width: 20px;
            }
        """)
        self.date_field.dateChanged.connect(self._update_preview)
        form_layout.addWidget(self.date_field)

        form_layout.addStretch()
        form_scroll.setWidget(form_container)
        self.input_section.set_content(form_scroll)
        main_layout.addWidget(self.input_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        text = self.generate_text()
        # Send directly to card
        self.sent.emit(text)

    def generate_text(self) -> str:
        lines = []

        name = self.name_field.text().strip()
        if name:
            lines.append(f"Signed: {name}")

        designation = self.designation_field.text().strip()
        if designation:
            lines.append(f"Designation: {designation}")

        qualifications = self.qualifications_field.text().strip()
        if qualifications:
            lines.append(f"Qualifications: {qualifications}")

        gmc = self.gmc_field.text().strip()
        if gmc:
            lines.append(f"Registration: {gmc}")

        date = self.date_field.date().toString("dd MMMM yyyy")
        lines.append(f"Date: {date}")

        return "\n".join(lines)

    def get_state(self) -> dict:
        return {
            "name": self.name_field.text(),
            "designation": self.designation_field.text(),
            "qualifications": self.qualifications_field.text(),
            "gmc": self.gmc_field.text(),
            "date": self.date_field.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        if not state:
            return

        self.name_field.setText(state.get("name", ""))
        self.designation_field.setText(state.get("designation", ""))
        self.qualifications_field.setText(state.get("qualifications", ""))
        self.gmc_field.setText(state.get("gmc", ""))

        date_str = state.get("date", "")
        if date_str:
            date = QDate.fromString(date_str, "yyyy-MM-dd")
            if date.isValid():
                self.date_field.setDate(date)

        self._update_preview()


# ================================================================
# GPR CIRCUMSTANCES POPUP (section 3 - with preview and yellow entries)
# ================================================================

class GPRCircumstancesPopup(QWidget):
    """Popup for Circumstances to this Admission with preview and yellow collapsible entries."""

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._entries = []
        self._extracted_checkboxes = []
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Main scroll area
        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ====================================================
        # SECTION 2: IMPORTED DATA (gold/yellow collapsible entries)
        # ====================================================
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(300)
        self.extracted_section._min_height = 100
        self.extracted_section._max_height = 500
        self.extracted_section.set_header_style("""
            QFrame {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-radius: 6px 6px 0 0;
            }
        """)
        self.extracted_section.set_title_style("""
            QLabel {
                font-size: 18px;
                font-weight: 600;
                color: #806000;
                background: transparent;
                border: none;
            }
        """)

        extracted_content = QWidget()
        extracted_content.setStyleSheet("""
            QWidget {
                background: rgba(255, 248, 220, 0.95);
                border: 1px solid rgba(180, 150, 50, 0.4);
                border-top: none;
                border-radius: 0 0 12px 12px;
            }
            QCheckBox {
                background: transparent;
                border: none;
                padding: 4px;
                font-size: 17px;
                color: #4a4a4a;
            }
            QCheckBox::indicator {
                width: 16px;
                height: 16px;
            }
        """)

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QScrollArea > QWidget > QWidget { background: transparent; }
        """)

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)
        self.extracted_checkboxes_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _update_preview(self):
        """Build text from selections and send directly to card."""
        parts = []
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                parts.append(cb.property("full_text"))

        combined = "\n\n".join(parts) if parts else ""
        # Send directly to card
        self.sent.emit(combined)

    def set_entries(self, items: list):
        """Display entries with collapsible dated entry boxes in yellow/gold UI."""
        self._entries = items

        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if isinstance(items, str):
            if items.strip():
                items = [{"date": None, "text": p.strip()} for p in items.split("\n\n") if p.strip()]
            else:
                items = []

        if not items:
            self.extracted_section.setVisible(False)
            return

        self.extracted_section.setVisible(True)

        # Sort by date (newest first)
        def get_sort_date(item):
            dt = item.get("date")
            if dt is None:
                return ""
            if hasattr(dt, "strftime"):
                return dt.strftime("%Y-%m-%d")
            return str(dt)

        sorted_items = sorted(items, key=get_sort_date, reverse=True)

        for item in sorted_items:
            dt = item.get("date")
            text = item.get("text", "").strip()
            if not text:
                continue

            if dt:
                if hasattr(dt, "strftime"):
                    date_str = dt.strftime("%d %b %Y")
                else:
                    date_str = str(dt)
            else:
                date_str = "No date"

            entry_frame = QFrame()
            entry_frame.setObjectName("entryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet("""
                QFrame#entryFrame {
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 8px;
                    padding: 4px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(180, 150, 50, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: #806000;
                }
                QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 17px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)
            header_row.addStretch()

            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(self._update_preview)
            header_row.addWidget(cb)

            entry_layout.addLayout(header_row)

            body_text = QTextEdit()
            body_text.setPlainText(text)
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet("""
                QTextEdit {
                    font-size: 17px;
                    color: #333;
                    background: rgba(255, 248, 220, 0.5);
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }
            """)
            body_text.setVisible(False)

            # Calculate height based on content
            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 400)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 200))

            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("▸")
                    else:
                        body.setVisible(True)
                        btn.setText("▾")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, btn=toggle_btn: btn.click()

            entry_layout.addWidget(body_text)
            self._extracted_checkboxes.append(cb)
            self.extracted_checkboxes_layout.addWidget(entry_frame)


# ================================================================
# FIXED DATA PANEL (for extracted data display)
# ================================================================

class GPRFixedDataPanel(QWidget):
    """Panel for displaying extracted data from documents."""

    sent = Signal(str)

    def __init__(self, title: str, subtitle: str = "", parent=None):
        super().__init__(parent)
        self._entries = []
        self._title = title
        self._subtitle = subtitle
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # Make the panel fill available space
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        title = QLabel(self._title)
        title.setStyleSheet("font-size: 20px; font-weight: 700; color: #2563eb;")
        layout.addWidget(title)

        if self._subtitle:
            subtitle = QLabel(self._subtitle)
            subtitle.setStyleSheet("font-size: 17px; color: #6b7280; font-style: italic;")
            layout.addWidget(subtitle)

        # Scrollable content - fills available space, no horizontal scroll
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                background: white;
            }
        """)

        self.content_widget = QWidget()
        self.content_widget.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(12, 12, 12, 12)
        self.content_layout.setSpacing(8)
        scroll.setWidget(self.content_widget)
        self._scroll = scroll
        self._scroll_height = 300
        layout.addWidget(scroll, 1)  # stretch factor 1 to fill space

        # Drag bar for resizing
        self.drag_bar = QFrame()
        self.drag_bar.setFixedHeight(10)
        self.drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.drag_bar.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 2px 40px;
            }
            QFrame:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        self.drag_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 300
        layout.addWidget(self.drag_bar)

        # Send all button
        send_btn = QPushButton("Send All to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #2563eb;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: 600;
            }
            QPushButton:hover { background: #1d4ed8; }
        """)
        send_btn.clicked.connect(self._send_all)
        layout.addWidget(send_btn)

    def set_entries(self, entries: list):
        """Set the entries to display."""
        self._entries = entries

        # Clear existing
        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        if not entries:
            no_data = QLabel("No data extracted for this section.")
            no_data.setStyleSheet("color: #9ca3af; font-style: italic;")
            self.content_layout.addWidget(no_data)
            return

        for entry in entries:
            frame = QFrame()
            frame.setStyleSheet("""
                QFrame {
                    background: #f9fafb;
                    border: 1px solid #e5e7eb;
                    border-radius: 6px;
                    padding: 8px;
                }
            """)
            frame_layout = QVBoxLayout(frame)
            frame_layout.setContentsMargins(8, 8, 8, 8)

            # Date if present
            if entry.get("date"):
                date_lbl = QLabel(str(entry["date"]))
                date_lbl.setStyleSheet("font-size: 16px; color: #6b7280;")
                frame_layout.addWidget(date_lbl)

            # Text content
            text = entry.get("text", "")
            text_lbl = QLabel(text[:500] + "..." if len(text) > 500 else text)
            text_lbl.setWordWrap(True)
            text_lbl.setStyleSheet("font-size: 17px; color: #374151;")
            frame_layout.addWidget(text_lbl)

            self.content_layout.addWidget(frame)

        self.content_layout.addStretch()

    def _send_all(self):
        """Send all entries as text."""
        if not self._entries:
            return
        parts = []
        for entry in self._entries:
            text = entry.get("text", "")
            if text:
                parts.append(text)
        self.sent.emit("\n\n".join(parts))

    def eventFilter(self, obj, event):
        """Handle drag events on the drag bar."""
        from PySide6.QtCore import QEvent
        if hasattr(self, 'drag_bar') and obj == self.drag_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._scroll_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(150, min(600, self._drag_start_height + delta))
                self._scroll_height = int(new_height)
                self._scroll.setFixedHeight(self._scroll_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def set_forensic_data(self, notes: list, extracted_entries: list = None):
        """Combine risk analysis (Physical Aggression, Property Damage, Sexual Behaviour) with data extractor forensic data.

        Display in date order with risk type badges, highlighted matches, and filter panel.
        """
        from risk_overview_panel import analyze_notes_for_risk
        from datetime import datetime
        import re
        import html

        print(f"[FORENSIC] set_forensic_data called with {len(notes) if notes else 0} notes, {len(extracted_entries) if extracted_entries else 0} extracted entries")

        all_incidents = []

        # Risk categories relevant to forensic history
        FORENSIC_RISK_CATEGORIES = ["Physical Aggression", "Property Damage", "Sexual Behaviour"]

        # Run risk analysis on notes
        if notes:
            results = analyze_notes_for_risk(notes)
            for cat_name in FORENSIC_RISK_CATEGORIES:
                cat_data = results.get("categories", {}).get(cat_name, {})
                for incident in cat_data.get("incidents", []):
                    all_incidents.append({
                        "date": incident.get("date"),
                        "text": incident.get("full_text", ""),
                        "matched": incident.get("matched", ""),
                        "subcategory": incident.get("subcategory", ""),
                        "severity": incident.get("severity", "medium"),
                        "category": cat_name,
                        "source": "notes",
                    })

        # Add data extractor entries
        if extracted_entries:
            for entry in extracted_entries:
                if isinstance(entry, dict):
                    text = entry.get("text", "")
                    date = entry.get("date")
                else:
                    text = str(entry)
                    date = None
                if text and text.strip():
                    all_incidents.append({
                        "date": date,
                        "text": text,
                        "matched": "",
                        "subcategory": "Data Extractor",
                        "severity": "medium",
                        "category": "Forensic History",
                        "source": "extractor",
                    })

        if not all_incidents:
            no_data = QLabel("No forensic data found.")
            no_data.setStyleSheet("color: #9ca3af; font-style: italic;")
            self.content_layout.addWidget(no_data)
            return

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                return datetime.min
            return d

        sorted_incidents = sorted(all_incidents, key=get_sort_date, reverse=True)

        # Store all incidents for filtering
        self._all_forensic_incidents = sorted_incidents
        self._current_filter = None

        # Store for send_all
        self._entries = [{"text": x["text"], "date": x.get("date")} for x in sorted_incidents]

        # Clear existing
        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Category colors
        self._cat_colors = {
            "Physical Aggression": "#b71c1c",
            "Property Damage": "#ff9800",
            "Sexual Behaviour": "#673ab7",
            "Forensic History": "#607d8b",
        }

        self._severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        # Collect unique labels (category or category:subcategory)
        labels = {}
        for inc in sorted_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat and subcat != "Data Extractor":
                label = f"{cat}: {subcat}"
            else:
                label = cat
            if label not in labels:
                labels[label] = self._cat_colors.get(cat, "#666666")

        # Create filter panel with horizontal scroll
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(40)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea {
                border: none;
                background: transparent;
            }
            QScrollBar:horizontal {
                height: 6px;
                background: #f0f0f0;
                border-radius: 3px;
            }
            QScrollBar::handle:horizontal {
                background: #c0c0c0;
                border-radius: 3px;
                min-width: 20px;
            }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 6px 12px;
                    font-size: 16px;
                    font-weight: 600;
                }}
                QPushButton:hover {{
                    background: {color}dd;
                }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text: self._apply_forensic_filter(lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row (hidden initially)
        self._filter_status_widget = QWidget()
        self._filter_status_widget.setStyleSheet("background: transparent;")
        self._filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(self._filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        self._filter_label = QLabel("Filtered by: ")
        self._filter_label.setStyleSheet("font-size: 17px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(self._filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 4px 10px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #dc2626;
            }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(self._remove_forensic_filter)
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(self._filter_status_widget)
        self.content_layout.addWidget(filter_container)

        # Container for incident entries (for re-rendering on filter)
        self._incidents_container = QWidget()
        self._incidents_container.setStyleSheet("background: transparent;")
        self._incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        self._incidents_layout = QVBoxLayout(self._incidents_container)
        self._incidents_layout.setContentsMargins(0, 0, 0, 0)
        self._incidents_layout.setSpacing(8)
        self.content_layout.addWidget(self._incidents_container)

        # Render all incidents
        self._render_forensic_incidents(sorted_incidents)

        self.content_layout.addStretch()

    def _apply_forensic_filter(self, label: str):
        """Apply filter to show only incidents matching the label."""
        self._current_filter = label
        self._filter_label.setText(f"Filtered by: {label}")
        self._filter_status_widget.setVisible(True)

        # Filter incidents
        filtered = []
        for inc in self._all_forensic_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat and subcat != "Data Extractor":
                inc_label = f"{cat}: {subcat}"
            else:
                inc_label = cat
            if inc_label == label:
                filtered.append(inc)

        self._render_forensic_incidents(filtered)

    def _remove_forensic_filter(self):
        """Remove filter and show all incidents."""
        self._current_filter = None
        self._filter_status_widget.setVisible(False)
        self._render_forensic_incidents(self._all_forensic_incidents)

    def _render_forensic_incidents(self, incidents: list):
        """Render the list of incidents."""
        import re
        import html
        from PySide6.QtWidgets import QTextEdit

        # Clear existing
        while self._incidents_layout.count():
            child = self._incidents_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        for incident in incidents:
            date = incident["date"]
            cat_name = incident["category"]
            text = incident["text"]
            matched = incident["matched"]
            subcat_name = incident["subcategory"]
            severity = incident["severity"]

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Get colors
            cat_color = self._cat_colors.get(cat_name, "#666666")
            sev_color = self._severity_colors.get(severity, "#666666")

            # Create HTML with highlighted matched text
            escaped_text = html.escape(text)
            if matched:
                escaped_matched = html.escape(matched)
                try:
                    pattern = re.compile(re.escape(escaped_matched), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_matched}</span>',
                        escaped_text
                    )
                except:
                    highlighted_html = escaped_text
            else:
                highlighted_html = escaped_text

            full_html = f'''
            <html>
            <body style="font-family: -apple-system, BlinkMacSystemFont, sans-serif; font-size: 17px; color: #333; margin: 0; padding: 0;">
            {highlighted_html}
            </body>
            </html>
            '''

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("forensicEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#forensicEntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid #e5e7eb;
                    border-left: 4px solid {cat_color};
                    border-radius: 8px;
                    padding: 2px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 6, 6, 6)
            entry_layout.setSpacing(4)

            # Header row with category badge, subcategory, date, severity
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Category badge
            badge_text = f"{cat_name}: {subcat_name}" if subcat_name and subcat_name != "Data Extractor" else cat_name
            cat_badge = QLabel(badge_text)
            cat_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 14px;
                    font-weight: 600;
                    color: white;
                    background: {cat_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 6px;
                }}
            """)
            header_row.addWidget(cat_badge)

            # Severity badge
            sev_badge = QLabel(severity.upper())
            sev_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 13px;
                    font-weight: 700;
                    color: white;
                    background: {sev_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 5px;
                }}
            """)
            header_row.addWidget(sev_badge)

            # Date label
            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 500;
                    color: #6b7280;
                    background: transparent;
                    border: none;
                }
            """)
            header_row.addWidget(date_label)

            header_row.addStretch()
            entry_layout.addLayout(header_row)

            # Text content
            body_text = QTextEdit()
            body_text.setReadOnly(True)
            body_text.setHtml(full_html)
            body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            body_text.setMinimumHeight(60)
            body_text.setMaximumHeight(120)
            body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            body_text.setStyleSheet("""
                QTextEdit {
                    background: #f9fafb;
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                    font-size: 17px;
                    color: #374151;
                }
            """)
            entry_layout.addWidget(body_text)

            self._incidents_layout.addWidget(entry_frame)


# ================================================================
# MAIN PAGE
# ================================================================

class GeneralPsychReportPage(QWidget):
    """Main page for creating General Psychiatric Reports."""

    go_back = Signal()

    # Sections based on CPA Report Medical template
    SECTIONS = [
        ("1. Patient Details", "patient_details"),
        ("2. Report Based On", "report_based_on"),
        ("3. Circumstances to this Admission", "circumstances"),
        ("4. Background Information", "background"),
        ("5. Past Medical History", "medical_history"),
        ("6. Past Psychiatric History", "psych_history"),
        ("7. Risk", "risk"),
        ("8. History of Substance Use", "substance_use"),
        ("9. Forensic History", "forensic"),
        ("10. Medication", "medication"),
        ("11. Mental Disorder", "diagnosis"),
        ("12. Legal Criteria for Detention", "legal_criteria"),
        ("13. Strengths", "strengths"),
        ("14. Signature", "signature"),
    ]

    # Map extracted categories to sections (both canonical and display names)
    # Includes variations with/without colons, different capitalization, etc.
    CATEGORY_TO_SECTION = {
        # Canonical names (uppercase - both with underscores and spaces)
        "FRONT_PAGE": "patient_details",
        "FRONT PAGE": "patient_details",
        "REPORT_BASED_ON": "report_based_on",
        "REPORT BASED ON": "report_based_on",
        "HISTORY_OF_PRESENTING_COMPLAINT": "circumstances",
        "HISTORY OF PRESENTING COMPLAINT": "circumstances",
        "BACKGROUND_HISTORY": "background",
        "BACKGROUND HISTORY": "background",
        "PAST_PSYCH": "psych_history",
        "PAST PSYCH": "psych_history",
        "PHYSICAL_HEALTH": "medical_history",
        "PHYSICAL HEALTH": "medical_history",
        "RISK": "risk",
        # Tribunal report risk sections (17 & 18) map to GPR section 7 (risk)
        "risk_harm": "risk",
        "RISK_HARM": "risk",
        "Incidents of harm to self or others": "risk",
        "risk_property": "risk",
        "RISK_PROPERTY": "risk",
        "Incidents of property damage": "risk",
        "DRUGS_AND_ALCOHOL": "substance_use",
        "DRUGS AND ALCOHOL": "substance_use",
        "FORENSIC": "forensic",
        "MENTAL_STATE": "diagnosis",
        "MENTAL STATE": "diagnosis",
        "SUMMARY": "legal_criteria",
        "STRENGTHS": "strengths",
        "MEDICATION": "medication",
        "LEGAL_CRITERIA": "legal_criteria",
        "LEGAL CRITERIA": "legal_criteria",

        # Circumstances variations
        "History of Presenting Complaint": "circumstances",
        "Circumstances of Admission": "circumstances",
        "Circumstances to Admission": "circumstances",
        "Circumstances to this admission": "circumstances",
        "Circumstances to this Admission": "circumstances",
        "Circumstances": "circumstances",

        # Background variations (includes risk behavior in childhood)
        "Background History": "background",
        "Background Information": "background",
        "Background information": "background",
        "Personal History": "background",
        "Risk behavior in childhood": "background",
        "Risk Behavior in Childhood": "background",
        "Risk behaviour in childhood": "background",

        # Past Psychiatric History variations
        "Past Psychiatric History": "psych_history",
        "Past Psychiatric History: Admissions to Hospital": "psych_history",
        "Psychiatric History": "psych_history",

        # Medical History variations
        "Physical Health": "medical_history",
        "Past Medical History": "medical_history",
        "Medical History": "medical_history",

        # Risk variations
        "Risk": "risk",
        "Risk Assessment": "risk",

        # Substance Use variations
        "Drug and Alcohol History": "substance_use",
        "History of Substance Use": "substance_use",
        "History of substance use": "substance_use",
        "Drugs and Alcohol": "substance_use",
        "Substance Use": "substance_use",

        # Forensic variations
        "Forensic History": "forensic",
        "Forensic History and history of violent behavior": "forensic",
        "Forensic History and History of Violent Behavior": "forensic",

        # Medication variations
        "Medication": "medication",
        "Medications": "medication",
        "Medication History": "medication",

        # Diagnosis/Mental Disorder variations
        "Mental State": "diagnosis",
        "Mental State Examination": "diagnosis",
        "Mental Disorder": "diagnosis",
        "Diagnosis": "diagnosis",

        # Legal Criteria variations
        "Summary": "legal_criteria",
        "Legal Criteria": "legal_criteria",
        "Legal Criteria for Detention": "legal_criteria",
        "Plan": "legal_criteria",

        # Strengths variations
        "Strengths": "strengths",
        "strengths": "strengths",
        "Strengths or positive factors": "strengths",
        "Strengths or Positive Factors": "strengths",
        "Positive factors": "strengths",
        "Positive Factors": "strengths",
    }

    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self.cards = {}
        self.popups = {}
        self._selected_card_key = None
        self._my_details = self._load_my_details()
        self._current_gender = "Male"  # Default gender for pronoun generation
        self._current_age = 40  # Default age for slider limits

        # Store extracted data at page level
        self._extracted_raw_notes = []
        self._extracted_categories = {}

        # Bidirectional sync flag to prevent recursive updates
        self._syncing = False

        # Track last popup text for each section to preserve user additions
        self._last_popup_text = {}

        # Guard flags to prevent reprocessing on navigation
        self._data_processed_id = None
        self._notes_processed_id = None

        # Imported report data (from DOCX tribunal report parser)
        self._imported_report_data = {}

        self._setup_ui()

        # Connect to SharedDataStore for cross-report data sharing
        self._connect_shared_store()

    def _connect_shared_store(self):
        """Connect to SharedDataStore for cross-report data sharing."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.patient_info_changed.connect(self._on_patient_info_changed)
            shared_store.notes_changed.connect(self._on_notes_changed)
            shared_store.extracted_data_changed.connect(self._on_extracted_data_changed)
            if hasattr(shared_store, 'report_sections_changed'):
                shared_store.report_sections_changed.connect(self._on_shared_report_sections_changed)
            print("[GPR] Connected to SharedDataStore signals (patient_info, notes, extracted_data, report_sections)")

            # Check for existing data in shared store
            self._check_shared_store_for_existing_data()
        except Exception as e:
            print(f"[GPR] Could not connect to SharedDataStore: {e}")

    def _check_shared_store_for_existing_data(self):
        """Check SharedDataStore for existing data when page is created."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()

            # Check for existing patient info
            patient_info = shared_store.patient_info
            if patient_info and any(patient_info.values()):
                print(f"[GPR] Found existing patient info in SharedDataStore")
                self._fill_patient_details(patient_info)

            # Check for existing notes
            notes = shared_store.notes
            if notes:
                print(f"[GPR] Found {len(notes)} existing notes in SharedDataStore")
                self._extracted_raw_notes = notes

            # Check for existing extracted data
            extracted_data = shared_store.extracted_data
            if extracted_data:
                print(f"[GPR] Found existing extracted data in SharedDataStore")
                self._on_extracted_data_changed(extracted_data)

            # Check for existing report sections from another form
            if hasattr(shared_store, 'report_sections'):
                report_sections = shared_store.report_sections
                report_source = shared_store.get_report_source()
                if report_sections and report_source and report_source != "general_psychiatric":
                    print(f"[GPR] Found existing report sections from {report_source}")
                    self._on_shared_report_sections_changed(report_sections, report_source)
        except Exception as e:
            print(f"[GPR] Error checking shared store: {e}")

    def _has_report_data(self):
        """Check if report data has been imported (local or via SharedDataStore)."""
        if self._imported_report_data:
            return True
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            source = shared_store.get_report_source()
            if source and source != "general_psychiatric" and shared_store.report_sections:
                return True
        except Exception:
            pass
        return False

    def _on_patient_info_changed(self, patient_info: dict):
        """Handle patient info updates from SharedDataStore."""
        if patient_info and any(patient_info.values()):
            print(f"[GPR] Received patient info from SharedDataStore: {list(k for k,v in patient_info.items() if v)}")
            self._fill_patient_details(patient_info)

    def _on_notes_changed(self, notes: list):
        """Handle notes updates from SharedDataStore."""
        if self._has_report_data():
            print("[GPR] Skipping _on_notes_changed - report data takes priority")
            return
        if notes:
            print(f"[GPR] Received {len(notes)} notes from SharedDataStore")
            self._extracted_raw_notes = notes
            # Refresh popups that use notes
            self._refresh_notes_based_popups()

    def _on_extracted_data_changed(self, data: dict):
        """Handle extracted data updates from SharedDataStore - full popup population."""
        if self._has_report_data():
            print("[GPR] Skipping _on_extracted_data_changed - report data takes priority")
            return
        if not data:
            return

        print(f"[GPR] Received extracted data from SharedDataStore: {list(data.keys())}")

        # Get categories from the data structure
        categories = data.get("categories", data)  # Handle both formats
        if not categories:
            return

        # Store at page level
        self._extracted_categories = categories

        # Delete cached popups so they get recreated with new data
        popups_to_refresh = ['strengths', 'risk', 'forensic', 'psych_history', 'background',
                            'medical_history', 'substance_use', 'medication', 'diagnosis',
                            'legal_criteria', 'circumstances']
        for key in popups_to_refresh:
            if key in self.popups:
                old_popup = self.popups[key]
                self.popup_stack.removeWidget(old_popup)
                old_popup.deleteLater()
                del self.popups[key]
                print(f"[GPR] Deleted cached '{key}' popup for refresh")

        # Update remaining popups with new data
        for key, popup in list(self.popups.items()):
            if hasattr(popup, 'set_entries'):
                self._populate_popup_with_extracted_data(key, popup)

        print(f"[GPR] Refreshed popups with SharedDataStore data")

    def _on_shared_report_sections_changed(self, sections: dict, source_form: str):
        """Handle report sections imported from another form (cross-talk)."""
        if source_form == "general_psychiatric":
            return  # Skip own exports

        print(f"[GPR] Cross-talk received from {source_form}: {len(sections)} sections")

        # Store imported data for popups
        for key, content in sections.items():
            if not content:
                continue
            # Map incoming key to GPR section key if needed
            gpr_key = self.CATEGORY_TO_SECTION.get(key, key)
            self._imported_report_data[gpr_key] = content

            # Delete cached popup for refresh
            if gpr_key in self.popups:
                old = self.popups.pop(gpr_key)
                self.popup_stack.removeWidget(old)
                old.deleteLater()

        print(f"[GPR] Cross-talk stored {len(sections)} sections from {source_form}")

    def _refresh_notes_based_popups(self):
        """Refresh popups that depend on notes data."""
        # Delete cached popups that use notes analysis
        notes_popups = ['risk', 'psych_history', 'substance_use', 'circumstances']
        for key in notes_popups:
            if key in self.popups:
                old_popup = self.popups[key]
                self.popup_stack.removeWidget(old_popup)
                old_popup.deleteLater()
                del self.popups[key]
                print(f"[GPR] Deleted cached '{key}' popup for notes refresh")

    def _load_my_details(self) -> dict:
        """Load clinician details from database."""
        if not self.db:
            return {}

        details = self.db.get_clinician_details()
        if not details:
            return {}

        return {
            "full_name": details[1] or "",
            "role_title": details[2] or "",
            "discipline": details[3] or "",
            "registration_body": details[4] or "",
            "registration_number": details[5] or "",
            "phone": details[6] or "",
            "email": details[7] or "",
            "team_service": details[8] or "",
            "hospital_org": details[9] or "",
            "ward_department": details[10] or "",
            "signature_block": details[11] or "",
        }

    def set_notes(self, notes: list):
        """
        Set notes from shared data store.
        Called by MainWindow when notes are available from another section.
        """
        if not notes:
            return

        if self._has_report_data():
            print("[GPR] Skipping set_notes - report data takes priority")
            return

        # Skip if these exact notes were already processed
        notes_sig = (len(notes), id(notes))
        if self._notes_processed_id == notes_sig:
            print(f"[GeneralPsych] Skipping set_notes - notes already processed")
            return
        self._notes_processed_id = notes_sig

        # Store raw notes at page level for use in sections
        self._extracted_raw_notes = notes

        # If data extractor exists, update its notes too
        if hasattr(self, '_data_extractor') and self._data_extractor:
            if hasattr(self._data_extractor, 'set_notes'):
                self._data_extractor.set_notes(notes)

        print(f"[GeneralPsych] Received {len(notes)} notes from shared store")

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("""
            QFrame {
                background: #2563eb;
                border-bottom: 1px solid #1d4ed8;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        # Back button
        back_btn = QPushButton("< Back")
        back_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: white;
                border: 1px solid rgba(255,255,255,0.3);
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 500;
            }
            QPushButton:hover { background: rgba(255,255,255,0.1); }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("General Psychiatric Report")
        title.setStyleSheet("""
            font-size: 20px;
            font-weight: 700;
            color: white;
        """)
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Clear Report button
        clear_btn = QPushButton("Clear Report - Start New")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 16px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_report)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = GPRToolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        # Connect uploaded docs menu to SharedDataStore
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())
        # Track last active editor (persists when toolbar clicked)
        self._active_editor = None

        # Helper to get active editor
        def get_active_editor():
            return self._active_editor

        # Helper to safely call editor method
        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        # Connect formatting signals
        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(
                        self,
                        "Spell Check",
                        "No spelling errors found."
                    )

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area with splitter
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("""
            QScrollArea {
                background: #f3f4f6;
                border: none;
            }
        """)
        self.main_splitter.addWidget(self.cards_holder)

        self.editor_root = QWidget()
        self.editor_root.setStyleSheet("background: #f3f4f6;")
        self.editor_layout = QVBoxLayout(self.editor_root)
        self.editor_layout.setContentsMargins(32, 24, 32, 24)
        self.editor_layout.setSpacing(16)
        self.cards_holder.setWidget(self.editor_root)

        # Right: Panel
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(350)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame {
                background: rgba(245,245,245,0.98);
                border-left: 1px solid rgba(0,0,0,0.08);
            }
        """)
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 450])

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        # Header row with title and lock button
        self.panel_header = QWidget()
        self.panel_header.setStyleSheet("""
            background: #008C7E;
            border-radius: 8px;
            margin-bottom: 8px;
        """)
        header_layout = QHBoxLayout(self.panel_header)
        header_layout.setContentsMargins(16, 12, 16, 12)
        header_layout.setSpacing(8)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 22px;
            font-weight: 700;
            color: #ffffff;
            background: transparent;
        """)
        header_layout.addWidget(self.panel_title, 1)

        # Lock button in header
        self.header_lock_btn = QPushButton("Unlocked")
        self.header_lock_btn.setFixedSize(70, 26)
        self.header_lock_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.header_lock_btn.setToolTip("Click to lock this section")
        self.header_lock_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255, 255, 255, 0.3);
                border: 2px solid rgba(255, 255, 255, 0.8);
                border-radius: 13px;
                font-size: 13px;
                font-weight: 600;
                color: white;
            }
            QPushButton:hover {
                background: rgba(255, 255, 255, 0.5);
            }
        """)
        self.header_lock_btn.clicked.connect(self._toggle_current_popup_lock)
        self.header_lock_btn.hide()
        header_layout.addWidget(self.header_lock_btn)

        panel_layout.addWidget(self.panel_header)

        # Popup stack
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("background: white;")
        self.popup_stack.setMinimumHeight(200)

        # Add a placeholder widget
        placeholder = QLabel("Click a section card to see its popup")
        placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        placeholder.setStyleSheet("font-size: 18px; color: #6b7280; background: white;")
        self.popup_stack.addWidget(placeholder)

        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        # Create all cards
        self._create_cards()

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self  # Capture reference to self for closure

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _create_cards(self):
        """Create all section cards."""
        for title, key in self.SECTIONS:
            card = GPRCardWidget(title, key, parent=self.editor_root)
            # Hook up focus event to register this editor as active
            self._hook_editor_focus(card.editor)
            card.clicked.connect(self._on_card_clicked)
            # Connect card text changes to sync with popup
            card.editor.textChanged.connect(lambda k=key: self._on_card_text_changed(k))
            self.cards[key] = card
            self.editor_layout.addWidget(card)

        self.editor_layout.addStretch()

    def _toggle_current_popup_lock(self):
        """Toggle lock on the currently active popup."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'toggle_lock'):
            popup.toggle_lock()
            self._update_header_lock_button()

    def _update_header_lock_button(self):
        """Update header lock button to match current popup state."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'is_locked'):
            is_locked = popup.is_locked()
            if is_locked:
                self.header_lock_btn.setText("Locked")
                self.header_lock_btn.setToolTip("Click to unlock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(239, 68, 68, 0.5);
                        border: 2px solid #ef4444;
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: white;
                    }
                    QPushButton:hover { background: rgba(239, 68, 68, 0.7); }
                """)
            else:
                self.header_lock_btn.setText("Unlocked")
                self.header_lock_btn.setToolTip("Click to lock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(255, 255, 255, 0.3);
                        border: 2px solid rgba(255, 255, 255, 0.8);
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: white;
                    }
                    QPushButton:hover { background: rgba(255, 255, 255, 0.5); }
                """)
            self.header_lock_btn.show()
        else:
            self.header_lock_btn.hide()

    def _set_current_popup(self, popup):
        """Set the current active popup and update lock button."""
        self._current_popup = popup
        self._update_header_lock_button()

    def _on_card_clicked(self, key: str):
        """Handle card click - show appropriate popup."""
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        # Deselect previous
        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)

        # Select new
        self._selected_card_key = key
        if key in self.cards:
            self.cards[key].setSelected(True)

        # Get or create popup
        if key not in self.popups:
            popup = self._create_popup(key)
            if popup:
                self.popups[key] = popup
                self.popup_stack.addWidget(popup)

        if key in self.popups:
            popup = self.popups[key]
            self.popup_stack.setCurrentWidget(popup)
            self._set_current_popup(popup)

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details popup from extracted demographics."""
        if not patient_info:
            return

        # Create popup if it doesn't exist
        if "patient_details" not in self.popups:
            popup = GPRPatientDetailsPopup(parent=self)
            popup.set_clinician_details(self._my_details)
            popup.sent.connect(lambda text: self._on_popup_sent("patient_details", text))
            popup.gender_changed.connect(self._on_gender_changed)  # Update gender-sensitive popups
            popup.age_changed.connect(self._on_age_changed)  # Update age-sensitive popups (sliders)
            self.popups["patient_details"] = popup
            self.popup_stack.addWidget(popup)

        # Fill the popup fields
        popup = self.popups["patient_details"]
        if hasattr(popup, 'fill_patient_info'):
            popup.fill_patient_info(patient_info)
            print(f"[GPRReport] Updated patient_details popup with demographics")

    def _get_circumstances_from_timeline(self) -> list:
        """
        Use timeline_builder to find the most recent admission and return
        notes from admission date to 2 weeks post admission.
        """
        from datetime import datetime, timedelta
        from timeline_builder import build_timeline

        # Get raw notes
        notes = getattr(self, '_extracted_raw_notes', [])
        if not notes:
            print("[GPR] No notes available for circumstances timeline")
            return []

        # Build timeline to find admissions
        episodes = build_timeline(notes)
        if not episodes:
            print("[GPR] No episodes found in timeline")
            return []

        # Find the most recent inpatient episode
        inpatient_episodes = [ep for ep in episodes if ep.get("type") == "inpatient"]
        if not inpatient_episodes:
            print("[GPR] No inpatient episodes found")
            return []

        # Get the most recent (last) inpatient episode
        most_recent = inpatient_episodes[-1]
        admission_start = most_recent["start"]
        admission_end = most_recent["end"]

        # Use actual admission end date from timeline
        if admission_end is None:
            admission_end = admission_start + timedelta(days=14)

        print(f"[GPR] Most recent admission: {admission_start} to {admission_end}")
        print(f"[GPR] Filtering notes from {admission_start} to {admission_end}")

        # Filter notes within the date range
        filtered_entries = []
        for note in notes:
            note_date = note.get("date")
            if not isinstance(note_date, datetime):
                continue

            note_date_only = note_date.date() if isinstance(note_date, datetime) else note_date
            # Ensure consistent date comparison (handle date vs datetime)
            start_date = admission_start.date() if isinstance(admission_start, datetime) else admission_start
            end_date = admission_end.date() if isinstance(admission_end, datetime) else admission_end
            if start_date <= note_date_only <= end_date:
                # Format as entry for GPRFixedDataPanel
                text = note.get("text") or note.get("content") or note.get("body") or ""
                if text.strip():
                    filtered_entries.append({
                        "date": note_date.strftime("%d/%m/%Y %H:%M") if note_date else "",
                        "text": text.strip(),
                        "type": note.get("type", ""),
                    })

        # Sort by date
        filtered_entries.sort(key=lambda x: x.get("date", ""))

        print(f"[GPR] Found {len(filtered_entries)} notes in circumstances date range")
        return filtered_entries

    def _create_popup(self, key: str) -> Optional[QWidget]:
        """Create the appropriate popup for a section."""
        if key == "patient_details":
            popup = GPRPatientDetailsPopup(parent=self)
            popup.set_clinician_details(self._my_details)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            popup.gender_changed.connect(self._on_gender_changed)  # Update gender-sensitive popups
            popup.age_changed.connect(self._on_age_changed)  # Update age-sensitive popups (sliders)
            return popup

        elif key == "report_based_on":
            popup = GPRReportBasedOnPopup(parent=self)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Pre-populate if we have extracted data
            self._populate_popup_with_extracted_data(key, popup)
            return popup

        elif key == "psych_history":
            popup = GPRPsychHistoryPopup(parent=self)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Pre-populate if we have extracted data
            self._populate_popup_with_extracted_data(key, popup)
            # Run timeline analysis to detect admissions
            notes = getattr(self, '_extracted_raw_notes', [])
            if notes:
                popup.set_notes(notes)
            return popup

        elif key == "background":
            # Use full background popup with form fields, risk dropdown, and found data
            popup = GPRBackgroundPopup(parent=self, gender=self._current_gender)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Pre-populate if we have extracted data
            self._populate_popup_with_extracted_data(key, popup)
            return popup

        elif key == "medical_history":
            # Use medical history popup with collapsible sections and drag bars
            popup = GPRMedicalHistoryPopup(parent=self, gender=self._current_gender)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Pre-populate if we have extracted data
            self._populate_popup_with_extracted_data(key, popup)
            return popup

        elif key == "risk":
            # Use risk popup with Current and Historical risk sections
            popup = GPRRiskPopup(parent=self)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))

            # Get notes and check for report data (tribunal sections 17 & 18)
            notes = getattr(self, '_extracted_raw_notes', [])

            # Collect ALL risk-related entries from extracted categories
            # For reports: look for risk_harm (17), risk_property (18), or general RISK category
            report_risk_entries = []
            for cat_key, cat_data in self._extracted_categories.items():
                if isinstance(cat_data, dict):
                    cat_name = cat_data.get("name", str(cat_key)).lower()
                    key_lower = str(cat_key).lower()
                    # Check for tribunal report section keys (17=risk_harm, 18=risk_property) OR general risk
                    if key_lower in ["risk_harm", "risk_property", "risk"] or \
                       "harm to self" in cat_name or "property damage" in cat_name or \
                       "incidents of harm" in cat_name or cat_name == "risk" or \
                       "risk assessment" in cat_name:
                        items = cat_data.get("items", [])
                        report_risk_entries.extend(items)

            # Detect REPORT vs NOTES:
            # - REPORT: No raw notes (or very few) + has extracted categories
            # - NOTES: Has many raw notes with dates -> run risk analysis
            is_report = (len(notes) == 0 and len(self._extracted_categories) > 0) or \
                        (len(notes) < 5 and len(report_risk_entries) > 0)

            print(f"[GPR] Risk section: {len(notes)} notes, {len(report_risk_entries)} risk entries, is_report={is_report}")

            if is_report and report_risk_entries:
                # REPORT UPLOAD: Display imported report data with dates
                self._populate_risk_extracted_section(popup, None, entries=report_risk_entries)
                print(f"[GPR] Risk: populated with REPORT data ({len(report_risk_entries)} entries)")
            elif notes and len(notes) >= 5:
                # NOTES UPLOAD: Run risk analysis on notes (existing logic)
                popup.set_notes_for_risk_analysis(notes)
                print(f"[GPR] Risk: populated with NOTES analysis ({len(notes)} notes)")
            elif report_risk_entries:
                # Few notes but have extracted entries - treat as report
                self._populate_risk_extracted_section(popup, None, entries=report_risk_entries)
            else:
                # Fall back to extracted data if nothing else
                self._populate_popup_with_extracted_data(key, popup)

            return popup

        elif key == "substance_use":
            # Use substance use popup with collapsible sections and drag bars
            popup = GPRSubstanceUsePopup(parent=self, gender=self._current_gender)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Set patient age to limit age sliders
            if hasattr(popup, 'set_patient_age'):
                popup.set_patient_age(self._current_age)
                print(f"[GPR] Substance popup created with patient age: {self._current_age}")
            # Combine substance misuse from risk analysis with data extractor findings
            notes = getattr(self, '_extracted_raw_notes', [])
            # Get extracted entries for substance_use from data extractor
            extracted_entries = []
            print(f"[SUBSTANCE] ========== POPUP OPENED ==========")
            print(f"[SUBSTANCE] _extracted_categories has {len(self._extracted_categories)} categories")
            print(f"[SUBSTANCE] Category keys: {list(self._extracted_categories.keys())}")
            print(f"[SUBSTANCE] CATEGORY_TO_SECTION keys for substance_use: {[k for k,v in self.CATEGORY_TO_SECTION.items() if v == 'substance_use']}")
            for cat_key, cat_data in self._extracted_categories.items():
                print(f"[SUBSTANCE] >>> Checking key='{cat_key}'")
                if isinstance(cat_data, dict):
                    cat_name = cat_data.get("name", str(cat_key))
                    print(f"[SUBSTANCE]     cat_name='{cat_name}', dict keys={list(cat_data.keys())}")
                    section_key = self.CATEGORY_TO_SECTION.get(cat_name)
                    # Also try the key itself
                    if not section_key:
                        section_key = self.CATEGORY_TO_SECTION.get(str(cat_key))
                    print(f"[SUBSTANCE]     -> section '{section_key}' (need 'substance_use')")
                    if section_key == "substance_use":
                        items = cat_data.get("items", [])
                        print(f"[SUBSTANCE]     *** MATCH! Found {len(items)} items")
                        for i, item in enumerate(items[:3]):  # Show first 3 items
                            if isinstance(item, dict):
                                print(f"[SUBSTANCE]     Item {i}: keys={list(item.keys())}, text={item.get('text', '')[:80]}")
                            else:
                                print(f"[SUBSTANCE]     Item {i}: {str(item)[:80]}")
                        extracted_entries.extend(items)
                else:
                    print(f"[SUBSTANCE]     NOT a dict: type={type(cat_data)}")
            print(f"[SUBSTANCE] ========== RESULT: {len(extracted_entries)} extracted entries, {len(notes)} notes ==========")
            # Call substance analysis with notes and extracted data
            popup.set_notes_for_substance_analysis(notes, extracted_entries)
            return popup

        elif key == "medication":
            # Use medication popup with collapsible sections and drag bars
            popup = GPRMedicationPopup(parent=self)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Pre-populate if we have extracted data
            self._populate_popup_with_extracted_data(key, popup)
            # Auto-fill medications from notes (last 6 months)
            notes = getattr(self, '_extracted_raw_notes', [])
            if notes:
                popup.prefill_medications_from_notes(notes)
            return popup

        elif key == "diagnosis":
            # Use diagnosis popup with ICD-10 selection
            popup = GPRDiagnosisPopup(parent=self)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Auto-fill diagnoses - imported data takes precedence over notes
            notes = getattr(self, '_extracted_raw_notes', [])
            # Get imported diagnosis entries AND substance_use entries
            imported_entries = []
            substance_entries = []
            for cat_key, cat_data in self._extracted_categories.items():
                # Get category name and look up section
                cat_name = cat_data.get("name", cat_key) if isinstance(cat_data, dict) else str(cat_key)
                section_key = self.CATEGORY_TO_SECTION.get(cat_name)
                if not section_key:
                    section_key = self.CATEGORY_TO_SECTION.get(str(cat_key))
                if section_key == "diagnosis":
                    items = cat_data.get("items", []) if isinstance(cat_data, dict) else []
                    imported_entries.extend(items)
                elif section_key == "substance_use":
                    items = cat_data.get("items", []) if isinstance(cat_data, dict) else []
                    substance_entries.extend(items)
            # Filter entries to only include those with diagnosis evidence
            relevant_entries = []
            for entry in imported_entries:
                text = entry.get("text", "") if isinstance(entry, dict) else str(entry)
                if text and popup._has_diagnosis_evidence(text):
                    relevant_entries.append(entry)
            # Only show entries that have diagnosis-related content
            if relevant_entries:
                print(f"[GPR] Diagnosis popup: {len(relevant_entries)} of {len(imported_entries)} entries have diagnosis evidence")
                popup.set_entries(relevant_entries)
            # Prefill - imported data takes precedence over notes, pass substance data for intelligent M&BD detection
            if notes or imported_entries:
                popup.prefill_diagnoses_from_notes(notes, imported_entries, substance_entries)
            return popup

        elif key == "legal_criteria":
            # Use legal criteria popup with nature/degree/necessity sections
            gender = getattr(self, '_current_gender', 'Male')
            from icd10_dict import ICD10_DICT
            popup = GPRLegalCriteriaPopup(parent=self, gender=gender, icd10_dict=ICD10_DICT)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            # Pre-populate if we have extracted data
            self._populate_popup_with_extracted_data(key, popup)
            return popup

        elif key == "strengths":
            # Use tribunal StrengthsPopup (same as tribunal section 13)
            from tribunal_popups import StrengthsPopup
            gender = getattr(self, '_current_gender', 'Male')
            popup = StrengthsPopup(parent=self, gender=gender)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))

            # Get notes and check for report data
            notes = getattr(self, '_extracted_raw_notes', [])

            # DEBUG: Log what we have
            print(f"[GPR STRENGTHS] ========== POPUP CREATED ==========")
            print(f"[GPR STRENGTHS] _extracted_categories has {len(self._extracted_categories)} categories")
            print(f"[GPR STRENGTHS] Category keys: {list(self._extracted_categories.keys())}")
            print(f"[GPR STRENGTHS] Notes count: {len(notes)}")

            # Get strengths data from extracted categories
            strengths_entries = []
            for cat_key, cat_data in self._extracted_categories.items():
                print(f"[GPR STRENGTHS] Checking cat_key='{cat_key}'")
                if isinstance(cat_data, dict):
                    cat_name = cat_data.get("name", str(cat_key)).lower()
                    key_lower = str(cat_key).lower()
                    print(f"[GPR STRENGTHS]   cat_name='{cat_name}', key_lower='{key_lower}'")
                    if "strengths" in cat_name or "strengths" in key_lower or \
                       "positive factors" in cat_name:
                        items = cat_data.get("items", [])
                        print(f"[GPR STRENGTHS]   *** MATCH! Found {len(items)} items")
                        strengths_entries.extend(items)

            # Detect REPORT vs NOTES
            is_report = (len(notes) == 0 and len(self._extracted_categories) > 0) or \
                        (len(notes) < 5 and len(strengths_entries) > 0)
            print(f"[GPR STRENGTHS] is_report={is_report}, strengths_entries={len(strengths_entries)}")

            if is_report and strengths_entries:
                # REPORT UPLOAD: Add imported data section to popup (with dates)
                print(f"[GPR STRENGTHS] Adding {len(strengths_entries)} entries with dates")
                self._add_imported_data_to_popup(popup, key, entries=strengths_entries)
                print(f"[GPR STRENGTHS] Populated with REPORT data ({len(strengths_entries)} entries)")
            else:
                print(f"[GPR STRENGTHS] NOT populating: is_report={is_report}, entries={len(strengths_entries)}")

            return popup

        elif key == "signature":
            # Use signature popup with auto-load from mydetails
            popup = GPRSignaturePopup(parent=self, my_details=self._my_details)
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))
            return popup

        elif key == "circumstances":
            # Use TribunalProgressPopup with narrative generation for Section 3
            # Filter to last admission only for circumstances
            from tribunal_popups import TribunalProgressPopup
            popup = TribunalProgressPopup(parent=self, date_filter='last_admission')
            popup.sent.connect(lambda text: self._on_popup_sent(key, text))

            # Populate using timeline to find most recent admission
            entries = self._get_circumstances_from_timeline()
            if entries:
                popup.set_entries(entries)

            return popup

        elif key == "forensic":
            # Use ForensicHistoryPopup (same as tribunal section 5) with index offence field
            from forensic_history_popup import ForensicHistoryPopup
            gender = getattr(self, '_current_gender', 'Male')
            popup = ForensicHistoryPopup(parent=self, gender=gender, show_index_offence=True)
            popup.sent.connect(lambda text, state: self._on_popup_sent(key, text))

            # Get notes and extracted data
            notes = getattr(self, '_extracted_raw_notes', [])
            extracted_entries = []
            for cat_key, cat_data in self._extracted_categories.items():
                if isinstance(cat_data, dict):
                    cat_name = cat_data.get("name", str(cat_key))
                    section_key = self.CATEGORY_TO_SECTION.get(cat_name)
                    if not section_key:
                        section_key = self.CATEGORY_TO_SECTION.get(str(cat_key))
                    if section_key == "forensic":
                        items = cat_data.get("items", [])
                        extracted_entries.extend(items)

            # Detect REPORT vs NOTES:
            # - REPORT: Few/no raw notes, extracted entries have longer text content
            # - NOTES: Many raw notes with dates, used for risk analysis
            is_report = False
            if extracted_entries:
                # Report: no raw notes OR fewer entries with longer text
                avg_text_len = sum(len(e.get("text", "")) for e in extracted_entries) / len(extracted_entries)
                # If no notes at all, it's definitely a report
                # If has entries with avg text > 50 chars and fewer than 15 entries, likely a report
                is_report = (len(notes) == 0) or (len(extracted_entries) < 15 and avg_text_len > 50)

            print(f"[GPR] Forensic section: {len(notes)} notes, {len(extracted_entries)} extracted entries, is_report={is_report}")

            if is_report and extracted_entries:
                # REPORT UPLOAD: populate the imported data section with entries (including dates)
                self._populate_forensic_extracted_section(popup, None, entries=extracted_entries)
                print(f"[GPR] Forensic: populated with REPORT data ({len(extracted_entries)} entries)")
            elif notes:
                # NOTES UPLOAD: use forensic data method for notes analysis
                if hasattr(popup, 'set_forensic_data'):
                    popup.set_forensic_data(notes, extracted_entries)
                    print(f"[GPR] Forensic: populated with NOTES analysis ({len(notes)} notes)")
            elif extracted_entries:
                # Fallback: just show extracted entries
                if hasattr(popup, 'set_extracted_data'):
                    popup.set_extracted_data(extracted_entries)

            return popup

        return None

    def _populate_popup_with_extracted_data(self, key: str, popup):
        """Populate a popup with extracted data if available.

        Works with GPRFixedDataPanel and GPRPsychHistoryPopup (both have set_entries).
        """
        if not hasattr(popup, 'set_entries'):
            return

        # Find ALL matching categories for this section and combine
        all_items = []

        # Handle extracted data structure: {idx: {"name": "Category Name", "items": [...]}}
        for cat_key, cat_data in self._extracted_categories.items():
            if not isinstance(cat_data, dict):
                continue

            # Get the category name from inside the data structure
            cat_name = cat_data.get("name", str(cat_key))

            # Look up which section this category maps to
            section_key = self.CATEGORY_TO_SECTION.get(cat_name)
            if not section_key:
                # Try the raw key as well (for backwards compatibility)
                section_key = self.CATEGORY_TO_SECTION.get(str(cat_key))

            if section_key == key:
                items = cat_data.get("items", [])
                all_items.extend(items)

        if all_items:
            print(f"[GPR] Populating popup '{key}' with {len(all_items)} items")
            popup.set_entries(all_items)

    def _on_popup_sent(self, key: str, text: str):
        """Handle when popup sends text to card - preserves user additions."""
        if key in self.cards:
            self._syncing = True
            try:
                # Sections that use prose format (not "Field: Value")
                prose_keys = (
                    "report_based_on",  # Checkbox format
                    "medical_history",  # Prose format from PhysicalHealthPopup
                    "background",       # Prose format from BackgroundHistoryPopup
                    "circumstances",    # Prose format
                    "psych_history",    # Prose format with sections
                    "risk",             # Prose format with CURRENT/HISTORICAL sections
                    "substance_use",    # Prose format from DrugsAlcoholPopup
                    "forensic",         # Prose format
                    "medication",       # Prose format from GPRMedicationPopup
                    "diagnosis",        # Prose format from GPRDiagnosisPopup
                    "legal_criteria",   # Prose format from GPRLegalCriteriaPopup
                    "strengths",        # Prose format
                )

                if key in prose_keys:
                    # For prose sections, preserve user additions
                    current_text = self.cards[key].editor.toPlainText()
                    last_popup_text = self._last_popup_text.get(key, "")

                    if not current_text.strip():
                        # Card is empty, just set the text
                        self.cards[key].editor.setPlainText(text)
                        self._last_popup_text[key] = text
                    elif not last_popup_text:
                        # First time popup sends - just set the text
                        self.cards[key].editor.setPlainText(text)
                        self._last_popup_text[key] = text
                    elif current_text.strip() == last_popup_text.strip():
                        # User hasn't added anything, safe to replace
                        self.cards[key].editor.setPlainText(text)
                        self._last_popup_text[key] = text
                    else:
                        # User has modified the card - preserve their additions
                        # DON'T update last_popup_text so we can always find the original popup portion

                        # Try to find and replace the old popup text with new popup text
                        if last_popup_text.strip() in current_text:
                            new_text = current_text.replace(last_popup_text.strip(), text.strip(), 1)
                            self.cards[key].editor.setPlainText(new_text)
                            self._last_popup_text[key] = text
                        else:
                            # Old popup text not found - maybe user modified it
                            # Append new text at the end to avoid losing anything
                            new_combined = current_text.rstrip() + "\n\n" + text.strip()
                            self.cards[key].editor.setPlainText(new_combined)
                            # Don't update last_popup_text - keep original for future comparisons
                else:
                    # For structured sections (Field: Value), use smart update
                    self._smart_update_card(key, text)
            finally:
                self._syncing = False

    def _on_gender_changed(self, gender: str):
        """Update gender-sensitive popups when gender changes in patient details."""
        self._current_gender = gender if gender else "Male"  # Default to Male if empty
        print(f"[GPR] Gender changed to: {self._current_gender}")

        # Update all gender-sensitive popups
        gender_sensitive_popups = ["background", "medical_history", "substance_use", "forensic", "legal_criteria"]
        for popup_key in gender_sensitive_popups:
            if popup_key in self.popups:
                popup = self.popups[popup_key]
                if hasattr(popup, 'update_gender'):
                    popup.update_gender(self._current_gender)
                    print(f"[GPR] Updated {popup_key} popup with gender: {self._current_gender}")

    def _on_age_changed(self, age: int):
        """Update age-sensitive popups when patient age changes."""
        self._current_age = age
        print(f"[GPR] Patient age changed to: {age}")

        # Update substance_use popup to limit age sliders
        if "substance_use" in self.popups:
            popup = self.popups["substance_use"]
            if hasattr(popup, 'set_patient_age'):
                popup.set_patient_age(age)
                print(f"[GPR] Updated substance_use popup with patient age: {age}")

    def _smart_update_card(self, key: str, new_text: str):
        """Update card text intelligently - update existing lines, add new lines, preserve custom content."""
        if key not in self.cards:
            return

        card = self.cards[key]
        current_text = card.editor.toPlainText()

        # If card is empty, just set the text
        if not current_text.strip():
            card.editor.setPlainText(new_text)
            return

        # Parse both texts into field dictionaries
        current_fields = self._parse_structured_text(current_text)
        new_fields = self._parse_structured_text(new_text)

        # Build updated text: update known fields, preserve unknown content
        result_lines = []
        used_fields = set()
        custom_lines = []

        # Process current text line by line
        for line in current_text.split('\n'):
            line_stripped = line.strip()
            if not line_stripped:
                result_lines.append(line)
                continue

            # Check if this line matches a known field pattern
            field_name = self._extract_field_name(line)
            if field_name and field_name in new_fields:
                # Update this field with new value
                result_lines.append(f"{field_name}: {new_fields[field_name]}")
                used_fields.add(field_name)
            elif field_name and field_name in current_fields:
                # Keep existing field value (not in new text)
                result_lines.append(line)
                used_fields.add(field_name)
            else:
                # Custom content - preserve it
                custom_lines.append(line)
                result_lines.append(line)

        # Add any new fields that weren't in current text
        for field_name, value in new_fields.items():
            if field_name not in used_fields:
                result_lines.append(f"{field_name}: {value}")

        card.editor.setPlainText('\n'.join(result_lines))

    def _parse_structured_text(self, text: str) -> dict:
        """Parse text with 'Field: Value' format into a dictionary."""
        fields = {}
        for line in text.split('\n'):
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    field_name = parts[0].strip()
                    value = parts[1].strip()
                    if field_name and value:
                        fields[field_name] = value
        return fields

    def _extract_field_name(self, line: str) -> str:
        """Extract field name from a 'Field: Value' line."""
        if ':' in line:
            parts = line.split(':', 1)
            if len(parts) == 2:
                return parts[0].strip()
        return None

    def _on_card_text_changed(self, key: str):
        """Handle card text changes - sync to popup if not already syncing."""
        if self._syncing:
            return

        if key not in self.cards:
            return

        card_text = self.cards[key].editor.toPlainText()

        # Parse card text and update corresponding popup
        if key in self.popups:
            self._syncing = True
            try:
                self._update_popup_from_card(key, card_text)
            finally:
                self._syncing = False

    def _update_popup_from_card(self, key: str, card_text: str):
        """Update popup fields from card text."""
        if key not in self.popups:
            return

        popup = self.popups[key]
        fields = self._parse_structured_text(card_text)

        if key == "patient_details" and hasattr(popup, 'name_edit'):
            # Patient details popup
            if 'Name' in fields:
                popup.name_edit.blockSignals(True)
                popup.name_edit.setText(fields['Name'])
                popup.name_edit.blockSignals(False)
            if 'Date of Birth' in fields:
                popup.dob_edit.blockSignals(True)
                date = QDate.fromString(fields['Date of Birth'], 'dd/MM/yyyy')
                if date.isValid():
                    popup.dob_edit.setDate(date)
                popup.dob_edit.blockSignals(False)
            if 'Section' in fields:
                popup.section_combo.blockSignals(True)
                popup.section_combo.setCurrentText(fields['Section'])
                popup.section_combo.blockSignals(False)
            if 'Admission Date' in fields:
                popup.admission_date_edit.blockSignals(True)
                date = QDate.fromString(fields['Admission Date'], 'dd/MM/yyyy')
                if date.isValid():
                    popup.admission_date_edit.setDate(date)
                popup.admission_date_edit.blockSignals(False)
            if 'Current Location' in fields:
                popup.location_edit.blockSignals(True)
                popup.location_edit.setText(fields['Current Location'])
                popup.location_edit.blockSignals(False)
            if 'Report By' in fields:
                popup.report_by_edit.blockSignals(True)
                popup.report_by_edit.setText(fields['Report By'])
                popup.report_by_edit.blockSignals(False)
            if 'Date Seen' in fields:
                popup.date_seen_edit.blockSignals(True)
                date = QDate.fromString(fields['Date Seen'], 'dd/MM/yyyy')
                if date.isValid():
                    popup.date_seen_edit.setDate(date)
                popup.date_seen_edit.blockSignals(False)

        elif key == "signature" and hasattr(popup, 'name_field'):
            # Signature popup - field names: Signed, Designation, Qualifications, Registration, Date
            if 'Signed' in fields:
                popup.name_field.blockSignals(True)
                popup.name_field.setText(fields['Signed'])
                popup.name_field.blockSignals(False)
            if 'Designation' in fields:
                popup.designation_field.blockSignals(True)
                popup.designation_field.setText(fields['Designation'])
                popup.designation_field.blockSignals(False)
            if 'Qualifications' in fields:
                popup.qualifications_field.blockSignals(True)
                popup.qualifications_field.setText(fields['Qualifications'])
                popup.qualifications_field.blockSignals(False)
            if 'Registration' in fields:
                popup.gmc_field.blockSignals(True)
                popup.gmc_field.setText(fields['Registration'])
                popup.gmc_field.blockSignals(False)

    def _populate_forensic_extracted_section(self, popup, content, entries: list = None):
        """Populate forensic popup's built-in extracted_section with checkboxes for imported report text.

        Args:
            popup: The forensic popup widget
            content: Text content (used if entries not provided)
            entries: List of dicts with 'text' and optional 'date' keys (preferred)
        """
        from PySide6.QtWidgets import QLabel, QVBoxLayout, QHBoxLayout, QWidget, QCheckBox, QFrame, QPushButton, QTextEdit
        from PySide6.QtCore import Qt
        from datetime import datetime

        if not hasattr(popup, 'extracted_section') or not popup.extracted_section:
            print(f"[GPR] Forensic popup has no extracted_section")
            return

        # Handle both entries list and content string
        if entries:
            items = entries
        elif content and content.strip():
            # Legacy: split content into paragraphs
            if '\n\n' in content:
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            else:
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            items = [{"text": p, "date": None} for p in paragraphs]
        else:
            return

        try:
            # Clear existing content in the extracted_checkboxes_layout
            if hasattr(popup, 'extracted_checkboxes_layout'):
                layout = popup.extracted_checkboxes_layout
                # Clear existing widgets
                while layout.count():
                    item = layout.takeAt(0)
                    if item.widget():
                        item.widget().deleteLater()

                # Sort by date (newest first) if dates available
                def get_sort_date(item):
                    if isinstance(item, dict):
                        d = item.get("date")
                        if d and isinstance(d, datetime):
                            return d
                    return datetime.min

                sorted_items = sorted(items, key=get_sort_date, reverse=True)
                popup._imported_checkboxes = []

                for item in sorted_items[:20]:  # Limit to 20 items
                    if isinstance(item, dict):
                        text = item.get("text", "")
                        date = item.get("date")
                    else:
                        text = str(item)
                        date = None

                    if not text or len(text.strip()) < 3:
                        continue

                    # Format date
                    if date:
                        if hasattr(date, "strftime"):
                            date_str = date.strftime("%d %b %Y")
                        else:
                            date_str = str(date)
                    else:
                        date_str = None

                    # Create entry frame with colored left border (matching other sections)
                    entry_frame = QFrame()
                    entry_frame.setObjectName("forensicImportFrame")
                    entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
                    entry_frame.setStyleSheet("""
                        QFrame#forensicImportFrame {
                            background: rgba(255, 255, 255, 0.95);
                            border: 1px solid rgba(180, 150, 50, 0.4);
                            border-left: 4px solid #607d8b;
                            border-radius: 8px;
                            padding: 2px;
                        }
                    """)
                    entry_layout_v = QVBoxLayout(entry_frame)
                    entry_layout_v.setContentsMargins(6, 6, 6, 6)
                    entry_layout_v.setSpacing(4)

                    # Header row: toggle → date → source badge → stretch → checkbox
                    header_row = QHBoxLayout()
                    header_row.setSpacing(8)

                    # Toggle button on the LEFT
                    toggle_btn = QPushButton("▸")
                    toggle_btn.setFixedSize(22, 22)
                    toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                    toggle_btn.setStyleSheet("""
                        QPushButton {
                            background: rgba(96, 125, 139, 0.2);
                            border: none;
                            border-radius: 4px;
                            font-size: 17px;
                            font-weight: bold;
                            color: #607d8b;
                        }
                        QPushButton:hover { background: rgba(96, 125, 139, 0.35); }
                    """)
                    header_row.addWidget(toggle_btn)

                    # Date label (if date available)
                    date_label = None
                    if date_str:
                        date_label = QLabel(f"📅 {date_str}")
                        date_label.setStyleSheet("""
                            QLabel {
                                font-size: 16px;
                                font-weight: 500;
                                color: #806000;
                                background: transparent;
                                border: none;
                            }
                        """)
                        date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                        header_row.addWidget(date_label)

                    # Source badge
                    source_badge = QLabel("Forensic")
                    source_badge.setStyleSheet("""
                        QLabel {
                            font-size: 14px;
                            font-weight: 600;
                            color: white;
                            background: #607d8b;
                            border: none;
                            border-radius: 3px;
                            padding: 2px 6px;
                        }
                    """)
                    header_row.addWidget(source_badge)
                    header_row.addStretch()

                    # Checkbox on the RIGHT
                    cb = QCheckBox()
                    cb.setProperty("full_text", text)
                    cb.setFixedSize(18, 18)
                    cb.setStyleSheet("""
                        QCheckBox { background: transparent; }
                        QCheckBox::indicator { width: 16px; height: 16px; }
                    """)
                    header_row.addWidget(cb)

                    entry_layout_v.addLayout(header_row)

                    # Body (full text, hidden by default)
                    body_text = QTextEdit()
                    body_text.setPlainText(text)
                    body_text.setReadOnly(True)
                    body_text.setFrameShape(QFrame.Shape.NoFrame)
                    body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
                    body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                    body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                    body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
                    body_text.setStyleSheet("""
                        QTextEdit {
                            font-size: 14px;
                            color: #333;
                            background: rgba(255, 248, 220, 0.5);
                            border: none;
                            padding: 6px;
                            border-radius: 4px;
                        }
                    """)
                    body_text.setMinimumHeight(50)
                    body_text.setMaximumHeight(150)
                    body_text.setVisible(False)
                    entry_layout_v.addWidget(body_text)

                    # Toggle function
                    def make_toggle(btn, body, frame, popup_self):
                        def toggle():
                            is_visible = body.isVisible()
                            body.setVisible(not is_visible)
                            btn.setText("▾" if not is_visible else "▸")
                            frame.updateGeometry()
                            if hasattr(popup_self, 'extracted_container'):
                                popup_self.extracted_container.updateGeometry()
                                popup_self.extracted_container.update()
                        return toggle

                    toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, popup)
                    toggle_btn.clicked.connect(toggle_fn)
                    if date_label:
                        date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                    # Connect checkbox to add/remove text from card
                    def make_handler(checkbox, txt):
                        def handler(checked):
                            if checked and "forensic" in self.cards:
                                current = self.cards["forensic"].editor.toPlainText()
                                if txt not in current:
                                    if current and not current.endswith('\n'):
                                        current += '\n\n'
                                    self.cards["forensic"].editor.setPlainText(current + txt)
                            elif not checked and "forensic" in self.cards:
                                current = self.cards["forensic"].editor.toPlainText()
                                if txt in current:
                                    new_text = current.replace(txt, '')
                                    while '\n\n\n' in new_text:
                                        new_text = new_text.replace('\n\n\n', '\n\n')
                                    self.cards["forensic"].editor.setPlainText(new_text.strip())
                        return handler

                    cb.toggled.connect(make_handler(cb, text))
                    popup._imported_checkboxes.append(cb)
                    layout.addWidget(entry_frame)

            # Show the section and expand it to show content
            popup.extracted_section.setVisible(True)
            # Set a larger content height if method available
            if hasattr(popup.extracted_section, 'set_content_height'):
                popup.extracted_section.set_content_height(300)
            popup._imported_data_added = True
            print(f"[GPR] Populated forensic extracted_section with {len(sorted_items)} entries (with dates)")

        except Exception as e:
            print(f"[GPR] Failed to populate forensic extracted_section: {e}")
            import traceback
            traceback.print_exc()

    def _populate_risk_extracted_section(self, popup, content, entries: list = None):
        """Populate risk popup's built-in extracted_section with checkboxes for imported report text.

        Args:
            popup: The risk popup widget
            content: Text content (used if entries not provided)
            entries: List of dicts with 'text' and optional 'date' keys (preferred)
        """
        from PySide6.QtWidgets import QLabel, QVBoxLayout, QHBoxLayout, QWidget, QCheckBox, QFrame, QPushButton, QTextEdit
        from PySide6.QtCore import Qt
        from datetime import datetime

        if not hasattr(popup, 'extracted_section') or not popup.extracted_section:
            print(f"[GPR] Risk popup has no extracted_section")
            return

        # Handle both entries list and content string
        if entries:
            items = entries
        elif content and content.strip():
            # Legacy: split content into paragraphs
            if '\n\n' in content:
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            else:
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            items = [{"text": p, "date": None} for p in paragraphs]
        else:
            return

        try:
            # Clear existing content in the extracted_checkboxes_layout
            if hasattr(popup, 'extracted_checkboxes_layout'):
                layout = popup.extracted_checkboxes_layout
                # Clear existing widgets
                while layout.count():
                    item = layout.takeAt(0)
                    if item.widget():
                        item.widget().deleteLater()

                # Sort by date (newest first) if dates available
                def get_sort_date(item):
                    if isinstance(item, dict):
                        d = item.get("date")
                        if d and isinstance(d, datetime):
                            return d
                    return datetime.min

                sorted_items = sorted(items, key=get_sort_date, reverse=True)
                popup._imported_checkboxes = []

                for item in sorted_items[:20]:  # Limit to 20 items
                    if isinstance(item, dict):
                        text = item.get("text", "")
                        date = item.get("date")
                    else:
                        text = str(item)
                        date = None

                    if not text or len(text.strip()) < 3:
                        continue

                    # Format date
                    if date:
                        if hasattr(date, "strftime"):
                            date_str = date.strftime("%d %b %Y")
                        else:
                            date_str = str(date)
                    else:
                        date_str = None

                    # Create entry frame with colored left border (matching other sections)
                    entry_frame = QFrame()
                    entry_frame.setObjectName("riskImportFrame")
                    entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
                    entry_frame.setStyleSheet("""
                        QFrame#riskImportFrame {
                            background: rgba(255, 255, 255, 0.95);
                            border: 1px solid rgba(180, 150, 50, 0.4);
                            border-left: 4px solid #dc2626;
                            border-radius: 8px;
                            padding: 2px;
                        }
                    """)
                    entry_layout_v = QVBoxLayout(entry_frame)
                    entry_layout_v.setContentsMargins(6, 6, 6, 6)
                    entry_layout_v.setSpacing(4)

                    # Header row: toggle → date → source badge → stretch → checkbox
                    header_row = QHBoxLayout()
                    header_row.setSpacing(8)

                    # Toggle button on the LEFT
                    toggle_btn = QPushButton("▸")
                    toggle_btn.setFixedSize(22, 22)
                    toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                    toggle_btn.setStyleSheet("""
                        QPushButton {
                            background: rgba(220, 38, 38, 0.2);
                            border: none;
                            border-radius: 4px;
                            font-size: 17px;
                            font-weight: bold;
                            color: #dc2626;
                        }
                        QPushButton:hover { background: rgba(220, 38, 38, 0.35); }
                    """)
                    header_row.addWidget(toggle_btn)

                    # Date label (if date available)
                    date_label = None
                    if date_str:
                        date_label = QLabel(f"📅 {date_str}")
                        date_label.setStyleSheet("""
                            QLabel {
                                font-size: 16px;
                                font-weight: 500;
                                color: #806000;
                                background: transparent;
                                border: none;
                            }
                        """)
                        date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                        header_row.addWidget(date_label)

                    # Source badge
                    source_badge = QLabel("Risk")
                    source_badge.setStyleSheet("""
                        QLabel {
                            font-size: 14px;
                            font-weight: 600;
                            color: white;
                            background: #dc2626;
                            border: none;
                            border-radius: 3px;
                            padding: 2px 6px;
                        }
                    """)
                    header_row.addWidget(source_badge)
                    header_row.addStretch()

                    # Checkbox on the RIGHT
                    cb = QCheckBox()
                    cb.setProperty("full_text", text)
                    cb.setFixedSize(18, 18)
                    cb.setStyleSheet("""
                        QCheckBox { background: transparent; }
                        QCheckBox::indicator { width: 16px; height: 16px; }
                    """)
                    header_row.addWidget(cb)

                    entry_layout_v.addLayout(header_row)

                    # Body (full text, hidden by default)
                    body_text = QTextEdit()
                    body_text.setPlainText(text)
                    body_text.setReadOnly(True)
                    body_text.setFrameShape(QFrame.Shape.NoFrame)
                    body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
                    body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                    body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                    body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
                    body_text.setStyleSheet("""
                        QTextEdit {
                            font-size: 14px;
                            color: #333;
                            background: rgba(255, 248, 220, 0.5);
                            border: none;
                            padding: 6px;
                            border-radius: 4px;
                        }
                    """)
                    body_text.setMinimumHeight(50)
                    body_text.setMaximumHeight(150)
                    body_text.setVisible(False)
                    entry_layout_v.addWidget(body_text)

                    # Toggle function
                    def make_toggle(btn, body, frame, popup_self):
                        def toggle():
                            is_visible = body.isVisible()
                            body.setVisible(not is_visible)
                            btn.setText("▾" if not is_visible else "▸")
                            frame.updateGeometry()
                            if hasattr(popup_self, 'extracted_container'):
                                popup_self.extracted_container.updateGeometry()
                                popup_self.extracted_container.update()
                        return toggle

                    toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, popup)
                    toggle_btn.clicked.connect(toggle_fn)
                    if date_label:
                        date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                    # Connect checkbox to add/remove text from card
                    def make_handler(checkbox, txt):
                        def handler(checked):
                            if checked and "risk" in self.cards:
                                current = self.cards["risk"].editor.toPlainText()
                                if txt not in current:
                                    if current and not current.endswith('\n'):
                                        current += '\n\n'
                                    self.cards["risk"].editor.setPlainText(current + txt)
                            elif not checked and "risk" in self.cards:
                                current = self.cards["risk"].editor.toPlainText()
                                if txt in current:
                                    new_text = current.replace(txt, '')
                                    while '\n\n\n' in new_text:
                                        new_text = new_text.replace('\n\n\n', '\n\n')
                                    self.cards["risk"].editor.setPlainText(new_text.strip())
                        return handler

                    cb.toggled.connect(make_handler(cb, text))
                    popup._imported_checkboxes.append(cb)
                    layout.addWidget(entry_frame)

            # Show the section and expand it to show content
            popup.extracted_section.setVisible(True)
            # Set a larger content height if method available
            if hasattr(popup.extracted_section, 'set_content_height'):
                popup.extracted_section.set_content_height(300)
            popup._imported_data_added = True
            print(f"[GPR] Populated risk extracted_section with {len(sorted_items)} entries (with dates)")

        except Exception as e:
            print(f"[GPR] Failed to populate risk extracted_section: {e}")
            import traceback
            traceback.print_exc()

    def _add_imported_data_to_popup(self, popup, section_key: str, content: str = None, entries: list = None):
        """Add imported data collapsible section to popup (for report uploads).

        Args:
            popup: The popup widget
            section_key: The section key for this popup
            content: Text content (used if entries not provided)
            entries: List of dicts with 'text' and optional 'date' keys (preferred)
        """
        from PySide6.QtWidgets import QLabel, QVBoxLayout, QHBoxLayout, QWidget, QCheckBox, QFrame, QPushButton, QTextEdit
        from PySide6.QtCore import Qt
        from background_history_popup import CollapsibleSection
        from datetime import datetime

        # Handle both entries list and content string
        if entries:
            items = entries
        elif content and content.strip():
            # Skip pointless content
            cleaned = content.strip()
            if cleaned.lower() in ('yes', 'no', 'n/a', 'yes.', 'no.'):
                return
            # Legacy: split content into paragraphs
            if '\n\n' in content:
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            else:
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
            items = [{"text": p, "date": None} for p in paragraphs]
        else:
            return

        try:
            # Find the popup's scroll_layout (StrengthsPopup has this)
            target_layout = None
            if hasattr(popup, 'scroll_layout'):
                target_layout = popup.scroll_layout
            elif hasattr(popup, 'layout') and callable(popup.layout):
                target_layout = popup.layout()

            if not target_layout:
                print(f"[GPR] No target layout found for '{section_key}'")
                return

            # Sort by date (newest first) if dates available
            def get_sort_date(item):
                if isinstance(item, dict):
                    d = item.get("date")
                    if d and isinstance(d, datetime):
                        return d
                return datetime.min

            sorted_items = sorted(items, key=get_sort_date, reverse=True)

            # Create CollapsibleSection for imported data
            extracted_section = CollapsibleSection("Imported Data", start_collapsed=False)
            extracted_section.set_content_height(150)
            extracted_section._min_height = 80
            extracted_section._max_height = 400
            extracted_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.5);
                    border-radius: 6px 6px 0 0;
                }
            """)
            extracted_section.title_label.setStyleSheet("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            # Create content widget with multiline checkbox entries
            extracted_content = QWidget()
            extracted_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)

            extracted_layout = QVBoxLayout(extracted_content)
            extracted_layout.setContentsMargins(12, 10, 12, 10)
            extracted_layout.setSpacing(8)

            # Get current card content for pre-checking
            card_text = ""
            if section_key in self.cards:
                card_text = self.cards[section_key].editor.toPlainText().lower()

            popup._imported_checkboxes = []

            # Section color mapping
            section_colors = {
                "diagnosis": "#9c27b0",
                "forensic": "#607d8b",
                "risk": "#dc2626",
                "strengths": "#22c55e",
                "background": "#3b82f6",
            }
            section_color = section_colors.get(section_key, "#806000")

            for item in sorted_items[:20]:  # Limit to 20 items
                if isinstance(item, dict):
                    text = item.get("text", "")
                    date = item.get("date")
                else:
                    text = str(item)
                    date = None

                if not text or len(text.strip()) < 3:
                    continue

                # Format date
                if date:
                    if hasattr(date, "strftime"):
                        date_str = date.strftime("%d %b %Y")
                    else:
                        date_str = str(date)
                else:
                    date_str = None

                # Create entry frame with colored left border (matching other sections)
                entry_frame = QFrame()
                entry_frame.setObjectName("importedEntryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
                entry_frame.setStyleSheet(f"""
                    QFrame#importedEntryFrame {{
                        background: rgba(255, 255, 255, 0.95);
                        border: 1px solid rgba(180, 150, 50, 0.4);
                        border-left: 4px solid {section_color};
                        border-radius: 8px;
                        padding: 2px;
                    }}
                """)
                entry_layout_v = QVBoxLayout(entry_frame)
                entry_layout_v.setContentsMargins(6, 6, 6, 6)
                entry_layout_v.setSpacing(4)

                # Header row: toggle → date → source badge → stretch → checkbox
                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                # Toggle button on the LEFT
                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet(f"""
                    QPushButton {{
                        background: rgba(128, 96, 0, 0.2);
                        border: none;
                        border-radius: 4px;
                        font-size: 17px;
                        font-weight: bold;
                        color: {section_color};
                    }}
                    QPushButton:hover {{ background: rgba(128, 96, 0, 0.35); }}
                """)
                header_row.addWidget(toggle_btn)

                # Date label (if date available)
                date_label = None
                if date_str:
                    date_label = QLabel(f"📅 {date_str}")
                    date_label.setStyleSheet("""
                        QLabel {
                            font-size: 16px;
                            font-weight: 500;
                            color: #806000;
                            background: transparent;
                            border: none;
                        }
                    """)
                    date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                    header_row.addWidget(date_label)

                # Source badge
                source_label = section_key.replace("_", " ").title()
                source_badge = QLabel(source_label)
                source_badge.setStyleSheet(f"""
                    QLabel {{
                        font-size: 14px;
                        font-weight: 600;
                        color: white;
                        background: {section_color};
                        border: none;
                        border-radius: 3px;
                        padding: 2px 6px;
                    }}
                """)
                header_row.addWidget(source_badge)
                header_row.addStretch()

                # Checkbox on the RIGHT
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("""
                    QCheckBox { background: transparent; }
                    QCheckBox::indicator { width: 16px; height: 16px; }
                """)

                # Pre-check if already in card
                if text.lower()[:50] in card_text:
                    cb.setChecked(True)

                header_row.addWidget(cb)
                entry_layout_v.addLayout(header_row)

                # Body (full text, hidden by default)
                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
                body_text.setStyleSheet("""
                    QTextEdit {
                        font-size: 14px;
                        color: #333;
                        background: rgba(255, 248, 220, 0.5);
                        border: none;
                        padding: 6px;
                        border-radius: 4px;
                    }
                """)
                body_text.setMinimumHeight(50)
                body_text.setMaximumHeight(150)
                body_text.setVisible(False)
                entry_layout_v.addWidget(body_text)

                # Toggle function
                def make_toggle(btn, body, frame):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame)
                toggle_btn.clicked.connect(toggle_fn)
                if date_label:
                    date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                # Connect checkbox to add/remove text from card
                def make_handler(checkbox, txt, key):
                    def handler(checked):
                        if checked and key in self.cards:
                            current = self.cards[key].editor.toPlainText()
                            if txt not in current:
                                if current and not current.endswith('\n'):
                                    current += '\n\n'
                                self.cards[key].editor.setPlainText(current + txt)
                        elif not checked and key in self.cards:
                            current = self.cards[key].editor.toPlainText()
                            if txt in current:
                                new_text = current.replace(txt, '')
                                while '\n\n\n' in new_text:
                                    new_text = new_text.replace('\n\n\n', '\n\n')
                                self.cards[key].editor.setPlainText(new_text.strip())
                    return handler

                cb.toggled.connect(make_handler(cb, text, section_key))
                popup._imported_checkboxes.append(cb)
                extracted_layout.addWidget(entry_frame)

            extracted_section.set_content(extracted_content)

            # Insert before stretch if present
            if hasattr(target_layout, 'insertWidget'):
                # Find stretch position
                insert_pos = target_layout.count()
                for i in range(target_layout.count()):
                    item = target_layout.itemAt(i)
                    if item and item.spacerItem():
                        insert_pos = i
                        break
                target_layout.insertWidget(insert_pos, extracted_section)
            else:
                target_layout.addWidget(extracted_section)

            popup._imported_data_added = True
            print(f"[GPR] Added imported data to '{section_key}' popup ({len(sorted_items)} entries with dates)")

        except Exception as e:
            print(f"[GPR] Failed to add imported data to popup: {e}")
            import traceback
            traceback.print_exc()

    def _go_back(self):
        """Navigate back to reports page."""
        self.go_back.emit()

    def _clear_report(self):
        """Clear all report content."""
        reply = QMessageBox.question(
            self,
            "Clear Report",
            "Are you sure you want to clear all content and start a new report?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            # Clear all cards
            for card in self.cards.values():
                card.editor.clear()

            # Clear extracted data
            self._extracted_categories = {}
            self._extracted_raw_notes = []
            if hasattr(self, '_imported_report_data'):
                self._imported_report_data = {}
            if hasattr(self, '_imported_report_sections'):
                self._imported_report_sections = {}

            # Clear data extractor if it exists
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                if hasattr(self._data_extractor_overlay, 'clear_extraction'):
                    self._data_extractor_overlay.clear_extraction()

            # Clear popup memory/state
            self.popups.clear()

            # Reset panel to default
            self.panel_title.setText("Select a section")

            # Clear the popup stack (except data extractor)
            while self.popup_stack.count() > 0:
                widget = self.popup_stack.widget(0)
                self.popup_stack.removeWidget(widget)
                if widget != getattr(self, '_data_extractor_overlay', None):
                    widget.deleteLater()

            # Restore signature card from my details
            if self._my_details and "signature" in self.cards:
                popup = GPRSignaturePopup(parent=self, my_details=self._my_details)
                popup.hide()
                text = popup.generate_text()
                if text and text.strip():
                    self.cards["signature"].editor.setPlainText(text)
                popup.deleteLater()

            print("[GPR] Report cleared")

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            from shared_data_store import get_shared_store
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file — DOCX tribunal reports get dedicated parser."""
        if file_path.lower().endswith('.docx'):
            try:
                from gpr_report_parser import parse_gpr_report
                result = parse_gpr_report(file_path)
                if result and result.get("sections"):
                    self._populate_from_parsed_report(result)
                    return
                else:
                    print("[GPR] No sections found in DOCX, falling back to Data Extractor")
            except Exception as e:
                print(f"[GPR] DOCX parsing failed: {e}, falling back to Data Extractor")
                import traceback
                traceback.print_exc()
        # Fallback for PDF/DOC/RTF or empty DOCX
        self._send_to_data_extractor(file_path)

    def _populate_from_parsed_report(self, result: dict):
        """Populate GPR sections from parsed tribunal report DOCX."""
        from PySide6.QtWidgets import QMessageBox
        import os

        sections = result.get('sections', {})
        patient_info = result.get('patient_info')
        source_file = result.get('source_file', 'report')
        fmt = result.get('format', 'unknown')

        # Ask add/replace
        action = self._ask_import_action(source_file, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        # Store imported data
        merged_sections = {}
        for section_key, content in sections.items():
            if not content:
                continue
            if action == "add" and section_key in self._imported_report_data:
                self._imported_report_data[section_key] += '\n\n' + content
            else:
                self._imported_report_data[section_key] = content
            merged_sections[section_key] = self._imported_report_data[section_key]
            print(f"[GPR] Stored report section '{section_key}' (ready in popup)")

        # Fill patient details
        if patient_info:
            self._fill_patient_details(patient_info)
            # Push to shared store
            try:
                from shared_data_store import get_shared_store
                shared_store = get_shared_store()
                shared_store.set_patient_info(patient_info, source="gpr_report")
            except Exception as e:
                print(f"[GPR] Error pushing patient info to SharedDataStore: {e}")

        # Populate popups with imported data
        for section_key, content in merged_sections.items():
            if not content:
                continue

            # Delete cached popup so it gets recreated with new data
            if section_key in self.popups:
                old_popup = self.popups.pop(section_key)
                self.popup_stack.removeWidget(old_popup)
                old_popup.deleteLater()
                print(f"[GPR] Deleted cached '{section_key}' popup for refresh")

            # Create popup and populate with imported data
            popup = self._create_popup(section_key)
            if popup:
                self.popups[section_key] = popup
                self.popup_stack.addWidget(popup)
                self._add_imported_data_to_popup(popup, section_key, content=content)
                print(f"[GPR] Populated '{section_key}' popup with imported data")

        # Push to SharedDataStore for cross-talk
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.set_report_sections(merged_sections, source_form="general_psychiatric")
        except Exception as e:
            print(f"[GPR] Error pushing report sections to SharedDataStore: {e}")

        # Show success
        mapped_count = len(merged_sections)
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"Successfully {action_word.lower()} report from:\n{source_file}\n\n"
            f"Format: {fmt}\n"
            f"{action_word} {mapped_count} sections to popups.\n\n"
            f"Click each card to review and send the content."
        )
        print(f"[GPR] {action_word} {mapped_count} sections from DOCX ({fmt}) to popups")

    def _ask_import_action(self, source_filename: str, import_type: str = "notes") -> str:
        """Ask user whether to add to existing imported data or replace it.

        Returns: 'add', 'replace', or 'cancel'
        """
        from PySide6.QtWidgets import QMessageBox

        has_existing = (hasattr(self, '_extracted_raw_notes') and bool(self._extracted_raw_notes)) or \
                       (hasattr(self, '_imported_report_data') and bool(self._imported_report_data))

        if not has_existing:
            return "replace"

        msg = QMessageBox(self)
        msg.setWindowTitle("Import Data")
        msg.setText(
            f"Clinical notes have already been loaded.\n\n"
            f"Would you like to add these notes to the existing set, or replace them?"
        )
        add_btn = msg.addButton("Add to Existing", QMessageBox.AcceptRole)
        replace_btn = msg.addButton("Replace All", QMessageBox.DestructiveRole)
        cancel_btn = msg.addButton("Cancel", QMessageBox.RejectRole)
        msg.setDefaultButton(add_btn)
        msg.exec()

        clicked = msg.clickedButton()
        if clicked == cancel_btn:
            return "cancel"
        elif clicked == replace_btn:
            return "replace"
        return "add"

    def _send_to_data_extractor(self, file_path: str):
        """Send a file to the data extractor for processing."""
        print(f"[GPR] _send_to_data_extractor called with: {file_path}")

        # Open the data extractor overlay
        self._open_data_extractor_overlay()

        # Load the file directly into the data extractor
        print(f"[GPR] Checking data extractor: hasattr={hasattr(self, '_data_extractor_overlay')}, exists={getattr(self, '_data_extractor_overlay', None) is not None}")

        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            has_load_file = hasattr(self._data_extractor_overlay, 'load_file')
            print(f"[GPR] Data extractor has load_file method: {has_load_file}")

            if has_load_file:
                print(f"[GPR] Calling load_file with: {file_path}")
                self._data_extractor_overlay.load_file(file_path)
                print(f"[GPR] ✅ Sent file to data extractor: {file_path}")
            else:
                print(f"[GPR] ❌ load_file method not found on data extractor!")
        else:
            print(f"[GPR] ❌ Data extractor overlay not available!")

    def _open_data_extractor_overlay(self):
        """Create the data extractor (hidden) for background processing."""
        from data_extractor_popup import DataExtractorPopup

        # Create data extractor if not exists
        if not hasattr(self, '_data_extractor_overlay') or not self._data_extractor_overlay:
            self._data_extractor_overlay = DataExtractorPopup(parent=self)
            self._data_extractor_overlay.hide()

            # Prevent it from closing itself after extraction
            self._data_extractor_overlay.close = lambda: None

            # Connect the data extraction signal
            if hasattr(self._data_extractor_overlay, 'data_extracted'):
                self._data_extractor_overlay.data_extracted.connect(self._on_data_extracted)

    def _on_data_extracted(self, data: dict):
        """Handle extracted data from data extractor.

        Only populates 'Extracted from Notes' in popups - does NOT prefill cards directly.
        """
        if self._has_report_data():
            print("[GPR] Skipping _on_data_extracted - report data takes priority")
            return
        print(f"[GPR] Data extracted: {list(data.keys())}")
        cov = data.get("_coverage")
        if cov and cov.get("uncategorised", 0) > 0:
            print(f"[GPR] Warning: {cov['uncategorised']} paragraphs uncategorised "
                  f"({cov['categorised']}/{cov['total_paragraphs']} categorised)")

        # Skip if this exact data was already processed
        categories = data.get("categories", {})
        cat_keys = tuple(sorted(categories.keys())) if categories else ()
        cat_count = sum(len(v.get("items", [])) if isinstance(v, dict) else 0 for v in categories.values())
        content_sig = (cat_keys, cat_count)
        if self._data_processed_id == content_sig:
            print(f"[GPR] Skipping _on_data_extracted - data already processed")
            return
        self._data_processed_id = content_sig

        print(f"[GPR] Categories received: {list(categories.keys())}")

        # Get raw notes - first try local extractor, then fall back to SharedDataStore
        raw_notes = []
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            raw_notes = getattr(self._data_extractor_overlay, 'notes', [])

        # Fall back to SharedDataStore if no local notes
        if not raw_notes:
            try:
                from shared_data_store import get_shared_store
                shared_store = get_shared_store()
                if shared_store.has_notes():
                    raw_notes = shared_store.notes
                    print(f"[GPR] Got {len(raw_notes)} notes from SharedDataStore (global import)")
            except Exception as e:
                print(f"[GPR] Error getting notes from SharedDataStore: {e}")

        # Ask add/replace if existing notes
        action = self._ask_import_action("", "notes")
        if action == "cancel":
            return
        if action == "add":
            existing_notes = getattr(self, '_extracted_raw_notes', []) or []
            raw_notes = existing_notes + raw_notes
            existing_cats = getattr(self, '_extracted_categories', {}) or {}
            for cat_name, cat_data in categories.items():
                if cat_name in existing_cats and isinstance(existing_cats[cat_name], dict) and isinstance(cat_data, dict):
                    existing_items = existing_cats[cat_name].get("items", [])
                    new_items = cat_data.get("items", [])
                    existing_cats[cat_name]["items"] = existing_items + new_items
                else:
                    existing_cats[cat_name] = cat_data
            categories = existing_cats

        # Store at page level
        self._extracted_categories = categories
        self._extracted_raw_notes = raw_notes

        # Count how many sections have data
        populated_count = 0
        for cat_key, cat_data in categories.items():
            if not isinstance(cat_data, dict):
                continue

            # Get category name from inside the data structure
            cat_name = cat_data.get("name", str(cat_key))
            section_key = self.CATEGORY_TO_SECTION.get(cat_name)
            if not section_key:
                section_key = self.CATEGORY_TO_SECTION.get(str(cat_key))

            print(f"[GPR] Category '{cat_name}' (key={cat_key}) -> section '{section_key}'")

            if section_key:
                items = cat_data.get("items", [])
                if items:
                    populated_count += 1
                    print(f"[GPR] Found {len(items)} items for {section_key}")
            else:
                print(f"[GPR] WARNING: No mapping for category '{cat_name}'")

        print(f"[GPR] Total sections with data: {populated_count}")

        # Update popups with new data - data goes to 'Extracted from Notes' section only
        for key, popup in self.popups.items():
            if hasattr(popup, 'set_entries'):
                self._populate_popup_with_extracted_data(key, popup)

        # Get patient info from SharedDataStore and fill patient details
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            patient_info = shared_store.patient_info
            if patient_info and any(patient_info.values()):
                print(f"[GPR] Found patient info in SharedDataStore: {list(k for k,v in patient_info.items() if v)}")
                self._fill_patient_details(patient_info)
        except Exception as e:
            print(f"[GPR] Error getting patient info from SharedDataStore: {e}")

        # For report-style popups (strengths, risk, forensic), delete cached popup
        # so it gets recreated with fresh imported data next time it's opened
        report_style_popups = ['strengths', 'risk', 'forensic']
        for key in report_style_popups:
            if key in self.popups:
                old_popup = self.popups[key]
                self.popup_stack.removeWidget(old_popup)
                old_popup.deleteLater()
                del self.popups[key]
                print(f"[GPR] Deleted cached '{key}' popup - will recreate with new data")

    def _export_docx(self):
        """Export the report to DOCX format."""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Report",
            f"General_Psychiatric_Report_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)

            # Title
            title = doc.add_paragraph()
            title_run = title.add_run("PSYCHIATRIC REPORT")
            title_run.bold = True
            title_run.font.size = Pt(16)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # Get patient details for header table
            patient_text = self.cards["patient_details"].editor.toPlainText() if "patient_details" in self.cards else ""
            patient_data = self._parse_patient_details(patient_text)

            # Header table
            table = doc.add_table(rows=5, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            header_data = [
                ["NAME OF PATIENT", patient_data.get("name", ""), "ADMISSION DATE", patient_data.get("admission_date", "")],
                ["DATE OF BIRTH", patient_data.get("dob", ""), "CURRENT LOCATION", patient_data.get("location", "")],
                ["AGE", patient_data.get("age", ""), "SECTION", patient_data.get("section", "")],
                ["REPORT BY", patient_data.get("report_by", ""), "DATE SEEN", patient_data.get("date_seen", "")],
                ["", "", "", ""],
            ]

            for i, row_data in enumerate(header_data):
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    cell = row.cells[j]
                    cell.text = cell_text
                    if j in (0, 2):  # Label columns
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

            doc.add_paragraph()

            # Add each section
            for title, key in self.SECTIONS:
                if key == "patient_details":
                    continue  # Already in header table
                if key == "signature":
                    continue  # Handled separately below with signature image

                content = self.cards[key].editor.toPlainText().strip() if key in self.cards else ""

                # Section heading
                heading = doc.add_heading(title, level=2)

                # Content
                if content:
                    doc.add_paragraph(content)
                else:
                    doc.add_paragraph("[No content]")

                doc.add_paragraph()

            # Signature section with image
            heading = doc.add_heading("14. Signature", level=2)

            # Add signature image if exists
            sig_image_path = os.path.expanduser("~/MyPsychAdmin/signature.png")
            if os.path.exists(sig_image_path):
                sig_para = doc.add_paragraph()
                sig_run = sig_para.add_run()
                sig_run.add_picture(sig_image_path, width=Inches(2.0))

            # Add signature text (name, designation, GMC, etc.)
            sig_text = self.cards["signature"].editor.toPlainText().strip() if "signature" in self.cards else ""
            if sig_text:
                sig = doc.add_paragraph()
                sig_run = sig.add_run(sig_text)
            else:
                # Fallback to my_details
                sig = doc.add_paragraph()
                sig_run = sig.add_run(self._my_details.get("signature_block", "") or f"{self._my_details.get('full_name', '')}\n{self._my_details.get('role_title', '')}")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Report saved to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library not installed. Please install with: pip install python-docx")
        except Exception as e:
            QMessageBox.warning(self, "Export Error", f"Failed to export: {e}")

    def _parse_patient_details(self, text: str) -> dict:
        """Parse patient details text into a dictionary."""
        data = {}
        for line in text.split("\n"):
            if ":" in line:
                key, value = line.split(":", 1)
                key = key.strip().lower().replace(" ", "_")
                data[key] = value.strip()
        return data

    def get_report_data(self) -> dict:
        """Get all report data as a dictionary."""
        data = {}
        for title, key in self.SECTIONS:
            if key in self.cards:
                data[key] = self.cards[key].editor.toPlainText().strip()
        return data

    def set_report_data(self, data: dict):
        """Set report data from a dictionary."""
        for key, content in data.items():
            if key in self.cards:
                self.cards[key].editor.setPlainText(content)
