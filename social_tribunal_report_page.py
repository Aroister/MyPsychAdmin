# ================================================================
#  SOCIAL CIRCUMSTANCES TRIBUNAL REPORT PAGE — Social Tribunal Report Writer (T133)
# ================================================================

from __future__ import annotations

import re
import sys
from datetime import datetime, timedelta
from html import unescape

from PySide6.QtCore import Qt, Signal, QSize, QEvent
from PySide6.QtGui import QColor, QFontDatabase
from PySide6.QtWidgets import QGraphicsDropShadowEffect
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QSplitter, QStackedWidget, QTextEdit,
    QSizePolicy, QPushButton, QToolButton, QComboBox, QColorDialog, QMessageBox,
    QLineEdit, QApplication, QCheckBox
)
from ui_effects import GlowCardMixin
from mypsy_richtext_editor import MyPsychAdminRichTextEditor
from shared_widgets import create_zoom_row, add_lock_to_popup
from spell_check_textedit import enable_spell_check_on_textedit
from utils.resource_path import resource_path


# ================================================================
# SOCIAL TOOLBAR
# ================================================================

class SocialToolbar(QWidget):
    """Toolbar for the Social Circumstances Tribunal Report Page."""

    # Formatting signals
    set_font_family = Signal(str)
    set_font_size = Signal(int)

    toggle_bold = Signal()
    toggle_italic = Signal()
    toggle_underline = Signal()

    set_text_color = Signal(QColor)
    set_highlight_color = Signal(QColor)

    set_align_left = Signal()
    set_align_center = Signal()
    set_align_right = Signal()
    set_align_justify = Signal()

    bullet_list = Signal()
    numbered_list = Signal()

    indent = Signal()
    outdent = Signal()

    undo = Signal()
    redo = Signal()

    insert_date = Signal()
    insert_section_break = Signal()

    export_docx = Signal()
    check_spelling = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(80)
        self.setStyleSheet("""
            SocialToolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 18px;
                font-weight: 500;
            }
            QToolButton:hover {
                background: rgba(0,0,0,0.08);
            }
            QComboBox {
                background: rgba(255,255,255,0.85);
                color: #333333;
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 17px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #e0e0e0;
            }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(False)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        container.setMinimumWidth(1200)  # Force scrollbar when viewport is smaller
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 8, 2)
        layout.setSpacing(10)

        # Export button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(160, 42)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 18px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #1d4ed8;
            }
            QToolButton:pressed {
                background: #1e40af;
            }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Uploaded Docs button (dropdown menu)
        from PySide6.QtWidgets import QMenu
        import_btn = QToolButton()
        import_btn.setText("Uploaded Docs")
        import_btn.setFixedSize(160, 38)
        import_btn.setPopupMode(QToolButton.InstantPopup)
        self.upload_menu = QMenu()
        import_btn.setMenu(self.upload_menu)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 18px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #059669;
            }
            QToolButton:pressed {
                background: #047857;
            }
            QToolButton::menu-indicator { image: none; }
        """)
        layout.addWidget(import_btn)

        # ---------------------------------------------------------
        # FONT FAMILY
        # ---------------------------------------------------------
        self.font_combo = QComboBox()
        self.font_combo.setFixedWidth(180)

        families = QFontDatabase.families()
        if sys.platform == "win32":
            preferred = ["Segoe UI", "Calibri", "Cambria", "Arial", "Times New Roman"]
        else:
            preferred = ["Avenir Next", "Avenir", "SF Pro Text", "Helvetica Neue", "Helvetica"]

        added = set()
        for f in preferred:
            if f in families:
                self.font_combo.addItem(f)
                added.add(f)
        for f in families:
            if f not in added:
                self.font_combo.addItem(f)

        self.font_combo.currentTextChanged.connect(self.set_font_family.emit)
        layout.addWidget(self.font_combo)

        # ---------------------------------------------------------
        # FONT SIZE
        # ---------------------------------------------------------
        self.size_combo = QComboBox()
        self.size_combo.setFixedWidth(65)
        for sz in [8, 9, 10, 11, 12, 14, 16, 18, 20, 22]:
            self.size_combo.addItem(str(sz))
        self.size_combo.currentTextChanged.connect(lambda v: self.set_font_size.emit(int(v)))
        layout.addWidget(self.size_combo)

        # Simple button helper - fixed size to prevent compression
        def btn(label, slot):
            b = QToolButton()
            b.setText(label)
            b.setMinimumWidth(36)
            b.clicked.connect(slot)
            return b

        # ---------------------------------------------------------
        # BASIC STYLES
        # ---------------------------------------------------------
        layout.addWidget(btn("B", self.toggle_bold.emit))
        layout.addWidget(btn("I", self.toggle_italic.emit))
        layout.addWidget(btn("U", self.toggle_underline.emit))

        # ---------------------------------------------------------
        # COLORS
        # ---------------------------------------------------------
        layout.addWidget(btn("A", self._choose_text_color))
        layout.addWidget(btn("🖍", self._choose_highlight_color))

        # ---------------------------------------------------------
        # ALIGNMENT
        # ---------------------------------------------------------
        layout.addWidget(btn("L", self.set_align_left.emit))
        layout.addWidget(btn("C", self.set_align_center.emit))
        layout.addWidget(btn("R", self.set_align_right.emit))
        layout.addWidget(btn("J", self.set_align_justify.emit))

        # ---------------------------------------------------------
        # LISTS / INDENTATION
        # ---------------------------------------------------------
        layout.addWidget(btn("•", self.bullet_list.emit))
        layout.addWidget(btn("1.", self.numbered_list.emit))
        layout.addWidget(btn("→", self.indent.emit))
        layout.addWidget(btn("←", self.outdent.emit))

        # ---------------------------------------------------------
        # UNDO / REDO
        # ---------------------------------------------------------
        layout.addWidget(btn("⟲", self.undo.emit))
        layout.addWidget(btn("⟳", self.redo.emit))

        # ---------------------------------------------------------
        # INSERTS
        # ---------------------------------------------------------
        layout.addWidget(btn("Date", self.insert_date.emit))
        layout.addWidget(btn("Break", self.insert_section_break.emit))

        # ---------------------------------------------------------
        # SPELL CHECK
        # ---------------------------------------------------------
        spell_btn = QToolButton()
        spell_btn.setText("Spell Check")
        spell_btn.setFixedSize(120, 38)
        spell_btn.setStyleSheet("""
            QToolButton {
                background: #f59e0b;
                color: white;
                font-size: 16px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 6px 12px;
            }
            QToolButton:hover { background: #d97706; }
            QToolButton:pressed { background: #b45309; }
        """)
        spell_btn.setToolTip("Jump to next spelling error")
        spell_btn.clicked.connect(self.check_spelling.emit)
        layout.addWidget(spell_btn)

        # Finalize scroll area
        scroll.setWidget(container)
        outer_layout.addWidget(scroll)

    # -----------------------------
    # COLOR PICKERS
    # -----------------------------
    def _choose_text_color(self):
        col = QColorDialog.getColor(QColor("black"), self)
        if col.isValid():
            self.set_text_color.emit(col)

    def _choose_highlight_color(self):
        col = QColorDialog.getColor(QColor("yellow"), self)
        if col.isValid():
            self.set_highlight_color.emit(col)


# ================================================================
# SOCIAL CARD WIDGET
# ================================================================

class SocialCardWidget(QFrame):
    """A clickable card for a social circumstances report section."""

    clicked = Signal(str)  # Emits the section key

    STYLE_NORMAL = """
        SocialCardWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
            padding: 16px;
        }
        SocialCardWidget:hover {
            border-color: #f59e0b;
            background: #fffbeb;
        }
    """

    STYLE_SELECTED = """
        SocialCardWidget {
            background: #fef3c7;
            border: 2px solid #f59e0b;
            border-left: 4px solid #d97706;
            border-radius: 12px;
            padding: 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Title
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        layout.addWidget(title_lbl)

        # Editor (rich text with formatting support)
        self.editor = MyPsychAdminRichTextEditor()
        self.editor.setPlaceholderText("Click to edit...")
        self.editor.setReadOnly(False)
        self._editor_height = 100
        self.editor.setMinimumHeight(60)
        self.editor.setMaximumHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
                color: #374151;
            }
        """)
        editor_zoom = create_zoom_row(self.editor, base_size=16)
        layout.addLayout(editor_zoom)
        layout.addWidget(self.editor)

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(12)
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #f59e0b;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        """Handle drag events on the expand bar."""
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(60, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setMinimumHeight(self._editor_height)
                self.editor.setMaximumHeight(self._editor_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        # Only emit if not clicking on the editor or expand bar
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        """Set the selected state of the card."""
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        """Return whether the card is selected."""
        return self._selected


# ================================================================
# SOCIAL HEADING WIDGET (for Yes/No sections with no card content)
# ================================================================

class SocialHeadingWidget(QFrame):
    """A clickable heading for social sections that don't need an editor card."""

    clicked = Signal(str)  # Emits the section key

    STYLE_NORMAL = """
        SocialHeadingWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 8px 16px;
        }
        SocialHeadingWidget:hover {
            border-color: #f59e0b;
            background: #fffbeb;
        }
    """

    STYLE_SELECTED = """
        SocialHeadingWidget {
            background: #fef3c7;
            border: 2px solid #f59e0b;
            border-left: 4px solid #d97706;
            border-radius: 8px;
            padding: 8px 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.setMinimumHeight(50)

        # Add subtle shadow
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(12)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 25))
        self.setGraphicsEffect(shadow)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)

        # Title - with word wrap enabled
        self.title_lbl = QLabel(title)
        self.title_lbl.setWordWrap(True)
        self.title_lbl.setStyleSheet("""
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        self.title_lbl.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        layout.addWidget(self.title_lbl, 1)

        # Click indicator arrow
        arrow = QLabel("\u25B6")
        arrow.setStyleSheet("""
            font-size: 12px;
            color: #f59e0b;
            background: transparent;
        """)
        arrow.setFixedWidth(20)
        arrow.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        layout.addWidget(arrow)

        # Dummy editor for compatibility
        self.editor = _SocialDummyEditor()

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        return self._selected


class _SocialDummyEditor:
    """Minimal dummy editor for SocialHeadingWidget compatibility."""
    def toPlainText(self):
        return ""
    def setPlainText(self, text):
        pass
    def toHtml(self):
        return ""
    def setHtml(self, html):
        pass
    def clear(self):
        pass


# ================================================================
# FIXED DATA PANEL FOR SOCIAL
# ================================================================

class SocialFixedDataPanel(QWidget):
    """A panel for displaying fixed data from extraction."""

    sent = Signal(str)

    def __init__(self, title: str, subtitle: str = "", parent=None):
        super().__init__(parent)
        self._title = title
        self._entries = []
        self.notes = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        title_label = QLabel(title)
        title_label.setStyleSheet("""
            font-size: 17px;
            font-weight: 600;
            color: #b45309;
        """)
        layout.addWidget(title_label)

        if subtitle:
            subtitle_label = QLabel(subtitle)
            subtitle_label.setStyleSheet("font-size: 12px; color: #6b7280;")
            layout.addWidget(subtitle_label)

        self.date_info = QLabel("")
        self.date_info.setStyleSheet("font-size: 11px; color: #9ca3af; font-style: italic;")
        layout.addWidget(self.date_info)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(12, 12, 12, 12)
        self.content_layout.setSpacing(8)
        scroll.setWidget(self.content_widget)

        layout.addWidget(scroll, 1)

        # Summary frame
        self.summary_frame = QFrame()
        self.summary_frame.setStyleSheet("""
            QFrame {
                background: #fef3c7;
                border: 1px solid #f59e0b;
                border-radius: 8px;
                padding: 8px;
            }
        """)
        self.summary_frame.hide()
        summary_layout = QVBoxLayout(self.summary_frame)
        summary_layout.setContentsMargins(8, 8, 8, 8)

        summary_header = QHBoxLayout()
        summary_label = QLabel("Summary")
        summary_label.setStyleSheet("font-weight: 600; color: #92400e;")
        summary_header.addWidget(summary_label)
        summary_header.addStretch()

        copy_summary_btn = QPushButton("Copy")
        copy_summary_btn.setFixedWidth(50)
        copy_summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 4px 8px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover { background: #d97706; }
        """)
        copy_summary_btn.clicked.connect(self._copy_summary)
        summary_header.addWidget(copy_summary_btn)

        close_btn = QPushButton("X")
        close_btn.setFixedWidth(24)
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #92400e;
                border: none;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover { color: #78350f; background: rgba(0,0,0,0.1); border-radius: 4px; }
        """)
        close_btn.clicked.connect(lambda: self.summary_frame.hide())
        summary_header.addWidget(close_btn)

        summary_layout.addLayout(summary_header)

        self.summary_text = QTextEdit()
        self.summary_text.setMaximumHeight(150)
        self.summary_text.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #fcd34d;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        summary_layout.addWidget(self.summary_text)
        layout.addWidget(self.summary_frame)

        # Buttons
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        self.summary_btn = QPushButton("Summary")
        self.summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover { background: #d97706; }
        """)
        self.summary_btn.clicked.connect(self._generate_summary)
        btn_layout.addWidget(self.summary_btn)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(self._send_to_letter)
        btn_layout.addWidget(send_btn)

        layout.addLayout(btn_layout)

    def set_entries(self, entries: list, date_range_info: str = ""):
        """Set the entries to display."""
        self._entries = entries

        if date_range_info:
            self.date_info.setText(date_range_info)

        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # For incident panels, show summary directly
        if "harm" in self._title.lower() or "property" in self._title.lower():
            self._show_incident_summary_directly()
            return

        for entry in entries[:50]:
            entry_frame = QFrame()
            entry_frame.setStyleSheet("""
                QFrame {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 6px;
                    padding: 8px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(8, 6, 8, 6)
            entry_layout.setSpacing(4)

            date_val = entry.get('date') or entry.get('datetime') or 'Unknown date'
            # Convert datetime to string if needed
            if hasattr(date_val, 'strftime'):
                date_str = date_val.strftime('%d/%m/%Y')
            else:
                date_str = str(date_val)
            date_label = QLabel(date_str)
            date_label.setStyleSheet("font-size: 11px; color: #6b7280; font-weight: 500;")
            entry_layout.addWidget(date_label)

            content = entry.get('content', '') or entry.get('text', '')
            preview = content[:200] + "..." if len(content) > 200 else content
            content_label = QLabel(preview)
            content_label.setWordWrap(True)
            content_label.setStyleSheet("font-size: 12px; color: #374151;")
            entry_layout.addWidget(content_label)

            self.content_layout.addWidget(entry_frame)

        self.content_layout.addStretch()

    def _show_incident_summary_directly(self):
        """Show filtered incident summary directly."""
        date_to_entries = {}
        for entry in self._entries:
            date_val = entry.get('date') or entry.get('datetime') or 'Unknown date'
            content = entry.get('content', '') or entry.get('text', '')
            if not content.strip():
                continue

            # Convert datetime to string if needed
            if hasattr(date_val, 'strftime'):
                date_str = date_val.strftime('%d/%m/%Y')
            else:
                date_str = str(date_val)
                try:
                    for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%d %b %Y', '%d-%m-%Y']:
                        try:
                            parsed_date = datetime.strptime(date_str.strip()[:10], fmt)
                            date_str = parsed_date.strftime('%d/%m/%Y')
                            break
                        except:
                            continue
                except:
                    pass

            if date_str not in date_to_entries:
                date_to_entries[date_str] = []
            date_to_entries[date_str].append(content.strip())

        def should_filter_line(line: str) -> bool:
            line_lower = line.lower().strip()
            filter_prefixes = ['diagnosis:', 'positive behaviour', 'to self:', 'to others:', 'risk:', 'risks', 'self neglect:']
            for prefix in filter_prefixes:
                if line_lower.startswith(prefix):
                    return True
            filter_phrases = ['without incident', 'no evidence', 'nothing to indicate', 'risk of', 'risk to',
                            'medication for agitation', 'call police if', 'less agitation', 'less aggression',
                            'reduced agitation', 'reduced aggression', 'police and ambulance to be called if',
                            'police to be called if', '(agitation)', 'previous', 'threatened to walk out', 'can be aggressive']
            for phrase in filter_phrases:
                if phrase in line_lower:
                    return True
            import re
            if re.search(r'\bnil\b', line_lower) or re.search(r'\bnon\b', line_lower) or re.search(r'\bno\b', line_lower):
                return True
            return False

        seen_content = set()
        output_lines = []
        sorted_dates = sorted(date_to_entries.keys(), key=lambda d: datetime.strptime(d, '%d/%m/%Y') if '/' in d else datetime.min, reverse=True)

        for date_str in sorted_dates:
            entries = date_to_entries[date_str]
            filtered_entries = []
            for content in entries:
                lines = content.split('\n')
                filtered_lines = []
                for line in lines:
                    if line.strip() and not should_filter_line(line):
                        normalized = ' '.join(line.lower().split())
                        if normalized not in seen_content:
                            seen_content.add(normalized)
                            filtered_lines.append(line.strip())
                if filtered_lines:
                    filtered_entries.append(' '.join(filtered_lines))
            if filtered_entries:
                combined = ' | '.join(filtered_entries)
                output_lines.append(f"{date_str}: {combined}")

        count = len(output_lines)
        self.date_info.setText(f"{count} incident(s) after filtering")
        self.summary_btn.hide()

        summary_frame = QFrame()
        summary_frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px;
            }
        """)
        summary_layout = QVBoxLayout(summary_frame)

        summary_text = QTextEdit()
        summary_text.setPlainText('\n'.join(output_lines) if output_lines else "No incidents found after filtering.")
        summary_text.setStyleSheet("""
            QTextEdit {
                background: #fafafa;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                font-size: 12px;
                padding: 8px;
            }
        """)
        summary_text.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        summary_layout.addWidget(summary_text, 1)

        summary_frame.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.content_layout.addWidget(summary_frame, 1)
        self.notes = output_lines

    def _clear(self):
        self._entries = []
        self.notes = []
        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        self.summary_frame.hide()

    def _copy_summary(self):
        from PySide6.QtWidgets import QApplication
        clipboard = QApplication.clipboard()
        clipboard.setText(self.summary_text.toPlainText())

    def _generate_summary(self):
        """Generate smart summary based on panel type."""
        if not self._entries:
            self.summary_text.setPlainText("No data to summarize.")
            self.summary_frame.show()
            return

        # For Progress panel, use risk-based summary
        if "Progress" in self._title:
            summary_parts = self._generate_risk_based_summary()
            if summary_parts:
                self.summary_text.setPlainText('\n'.join(summary_parts))
                self.summary_frame.show()
                return

        # Default: simple summary
        summary_lines = []
        for entry in self._entries[:20]:
            date_val = entry.get('date') or entry.get('datetime') or ''
            if hasattr(date_val, 'strftime'):
                date_str = date_val.strftime('%d/%m/%Y')
            else:
                date_str = str(date_val)
            content = entry.get('content', '') or entry.get('text', '')
            preview = content[:100] + "..." if len(content) > 100 else content
            summary_lines.append(f"{date_str}: {preview}")

        self.summary_text.setPlainText('\n'.join(summary_lines))
        self.summary_frame.show()

    def _generate_risk_based_summary(self):
        """Generate narrative summary with Progress, Engagement, and Insight sections."""
        import re
        from datetime import datetime
        from pathlib import Path

        if not self._entries:
            return []

        # Load risk dictionary for scoring
        risk_dict = {}
        risk_file = Path(__file__).parent / "riskDICT.txt"
        if risk_file.exists():
            with open(risk_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or ',' not in line:
                        continue
                    parts = line.rsplit(',', 1)
                    term = parts[0].strip().lower()
                    score_str = parts[1].strip() if len(parts) > 1 else ''
                    if term and score_str:
                        try:
                            risk_dict[term] = int(score_str)
                        except:
                            pass

        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        def get_highest_scoring_line(content, risk_dict):
            """Return the line with the highest risk score (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            best_line = None
            best_score = -1
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                line_score = sum(points for term, points in risk_dict.items() if term in line_lower)
                if line_score > best_score:
                    best_score = line_score
                    best_line = cleaned
            return best_line

        def extract_specific_line(content, keywords):
            """Extract the specific line containing any of the keywords (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                if line_lower.endswith(':') and len(line_lower) < 40:
                    continue
                for kw in keywords:
                    if kw in line_lower:
                        for prefix in ['Observations and Engagement:', 'Engagement:', 'Self-care:', 'Self care:', 'Insight:']:
                            if cleaned.startswith(prefix):
                                cleaned = cleaned[len(prefix):].strip()
                        if len(cleaned) >= 100:
                            return cleaned
            return None

        # Categorize entries
        entries_data = []
        for entry in self._entries:
            content = entry.get('content', '') or entry.get('text', '') or ''
            content_lower = content.lower()
            if not content:
                continue

            note_date = parse_date(entry.get('date') or entry.get('datetime'))
            date_str = note_date.strftime('%d/%m/%Y') if note_date else 'Unknown'

            # Calculate score
            score = sum(points for term, points in risk_dict.items() if term in content_lower)

            entries_data.append({
                'date': note_date,
                'date_str': date_str,
                'score': score,
                'content': content,
                'content_lower': content_lower
            })

        # Sort by date
        entries_data.sort(key=lambda x: x['date'] or datetime.min)

        # === EXTRACT EVENTS BY CATEGORY ===
        concern_events = [e for e in entries_data if e['score'] >= 1500]

        violence_keywords = ['violence', 'violent', 'assault', 'attack', 'fight', 'aggression', 'aggressive', 'physical altercation']
        violence_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in violence_keywords)]

        positive_keywords = ['stable', 'settled', 'calm', 'pleasant', 'appropriate', 'well presented', 'good rapport', 'engaging well', 'cooperative', 'progress']
        positive_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in positive_keywords) and e['score'] < 100]

        selfcare_keywords = ['well kempt', 'good self care', 'good self-care', 'poor self care', 'poor self-care',
                            'well dressed', 'showered', 'washed', 'malodorous', 'smelt', 'clean', 'unkempt',
                            'dishevelled', 'neglected appearance', 'poor hygiene']
        selfcare_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in selfcare_keywords)]

        insight_keywords = ['insight', 'awareness', 'understands', 'accepts diagnosis', 'denies illness', 'lacks insight', 'good insight', 'poor insight', 'partial insight']
        insight_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in insight_keywords)]

        # === BUILD NARRATIVE SUMMARY ===
        summary_parts = []

        # Date range header
        dates = [e['date'] for e in entries_data if e['date']]
        if dates:
            earliest = min(dates).strftime('%d/%m/%Y')
            latest = max(dates).strftime('%d/%m/%Y')

        # === SECTION 1: PROGRESS ===
        summary_parts.append("PROGRESS")
        summary_parts.append("-" * 40)

        if dates:
            summary_parts.append(f"Review period: {earliest} to {latest}")

        # Concern events
        if concern_events:
            summary_parts.append("")
            summary_parts.append(f"Significant behavioural concerns noted ({len(concern_events)}):")
            for e in sorted(concern_events, key=lambda x: x['date'] or datetime.min, reverse=True):
                relevant_line = get_highest_scoring_line(e['content'], risk_dict)
                if relevant_line:
                    summary_parts.append(f"  * {e['date_str']}: {relevant_line}")
        else:
            summary_parts.append("")
            summary_parts.append("No significant behavioural concerns in this period.")

        # Violence/Aggression
        summary_parts.append("")
        summary_parts.append("Episodes of violence or aggression:")
        high_violence = [e for e in violence_events if e['score'] >= 2000]
        if high_violence:
            summary_parts.append(f"  ({len(high_violence)} episodes)")
            for e in sorted(high_violence, key=lambda x: x['date'] or datetime.min, reverse=True):
                relevant_line = get_highest_scoring_line(e['content'], risk_dict)
                if relevant_line:
                    summary_parts.append(f"  * {e['date_str']}: {relevant_line}")
        else:
            summary_parts.append("  Nil noted.")

        # Positive progress
        def summarize_positive(content_lower):
            if 'superficially settled' in content_lower or 'superfically settled' in content_lower:
                return 'settled (superficially)'
            if 'settled' in content_lower or 'calm' in content_lower:
                return 'settled'
            if 'stable' in content_lower:
                return 'stable'
            if 'pleasant' in content_lower:
                return 'pleasant'
            if 'cooperative' in content_lower:
                return 'cooperative'
            if 'engaging well' in content_lower or 'good rapport' in content_lower:
                return 'engaging well'
            if 'well presented' in content_lower or 'appropriate' in content_lower:
                return 'well presented'
            if 'progress' in content_lower:
                return 'progress noted'
            return 'positive'

        if positive_events:
            sorted_positive = sorted(positive_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append("")
            summary_parts.append(f"Positive observations ({len(sorted_positive)}):")
            for e in sorted_positive:
                descriptor = summarize_positive(e['content_lower'])
                summary_parts.append(f"  * {e['date_str']}: {descriptor}")

        # === SECTION 2: SELF-CARE ===
        summary_parts.append("")
        summary_parts.append("")
        summary_parts.append("SELF-CARE")
        summary_parts.append("-" * 40)
        if selfcare_events:
            sorted_selfcare = sorted(selfcare_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_selfcare)} records)")
            for e in sorted_selfcare:
                specific_line = extract_specific_line(e['content'], selfcare_keywords)
                if specific_line:
                    summary_parts.append(f"  * {e['date_str']}: {specific_line}")
        else:
            summary_parts.append("  No specific self-care records found.")

        # === SECTION 3: INSIGHT ===
        summary_parts.append("")
        summary_parts.append("")
        summary_parts.append("INSIGHT")
        summary_parts.append("-" * 40)
        if insight_events:
            sorted_insight = sorted(insight_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_insight)} records)")
            for e in sorted_insight:
                specific_line = extract_specific_line(e['content'], insight_keywords)
                if specific_line:
                    summary_parts.append(f"  * {e['date_str']}: {specific_line}")
        else:
            summary_parts.append("  No specific insight records found.")

        return summary_parts

    def _send_to_letter(self):
        if self.notes:
            self.sent.emit('\n'.join(self.notes))
        elif self._entries:
            content_parts = []
            for entry in self._entries[:10]:
                date_str = entry.get('date') or entry.get('datetime') or ''
                content = entry.get('content', '') or entry.get('text', '')
                if content:
                    content_parts.append(f"{date_str}: {content[:200]}")
            self.sent.emit('\n'.join(content_parts))


# ================================================================
# MAIN SOCIAL CIRCUMSTANCES TRIBUNAL REPORT PAGE
# ================================================================

class SocialTribunalReportPage(QWidget):
    """Main page for creating Social Circumstances Tribunal Reports (T133)."""

    go_back = Signal()

    # Sections based on T133 form - Social Circumstances Report
    SECTIONS = [
        ("1. Patient Details", "patient_details"),
        ("2. Factors affecting understanding or ability to cope with hearing", "factors_hearing"),
        ("3. Adjustments for tribunal to consider", "adjustments"),
        ("4. Index offence(s) and relevant forensic history", "forensic"),
        ("5. Previous involvement with mental health services", "previous_mh_dates"),
        ("6. Home and family circumstances", "home_family"),
        ("7. Housing or accommodation if discharged", "housing"),
        ("8. Financial position and benefit entitlements", "financial"),
        ("9. Employment opportunities if discharged", "employment"),
        ("10. Previous response to community support or Section 117 aftercare", "previous_community"),
        ("11. Care pathway and Section 117 after-care available", "care_pathway"),
        ("12. Proposed care plan", "care_plan"),
        ("13. Adequacy of proposed care plan", "care_plan_adequacy"),
        ("14. Funding issues for proposed care plan", "care_plan_funding"),
        ("15. Strengths or positive factors", "strengths"),
        ("16. Current progress, behaviour, compliance and insight", "progress"),
        ("17. Incidents of harm to self or others", "risk_harm"),
        ("18. Incidents of property damage", "risk_property"),
        ("19. Patient's views, wishes, beliefs, opinions, hopes and concerns", "patient_views"),
        ("20. Nearest Relative views", "nearest_relative"),
        ("21. Reasons if inappropriate to consult Nearest Relative", "nr_inappropriate"),
        ("22. Views of other person taking lead role in care (non-professional)", "carer_views"),
        ("23. MAPPA involvement", "mappa"),
        ("24. MCA 2005 deprivation of liberty consideration", "mca_dol"),
        ("25. Section 2: Detention justified for health, safety or protection", "s2_detention"),
        ("26. Other sections: Medical treatment justified", "other_detention"),
        ("27. Risk if discharged from hospital", "discharge_risk"),
        ("28. Community risk management", "community"),
        ("29. Other relevant information", "other_info"),
        ("30. Recommendations to tribunal", "recommendations"),
        ("31. Signature", "signature"),
    ]

    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self.cards = {}
        self.popups = {}
        self.popup_memory = {}
        self._active_editor = None
        self._current_gender = None  # Default to "they" pronouns
        self._selected_card_key = None
        self._my_details = self._load_my_details()

        self._extracted_raw_notes = []
        self._extracted_categories = {}
        self._incident_data = []

        # Guard flags to prevent reprocessing on navigation
        self._data_processed_id = None
        self._notes_processed_id = None

        self._setup_ui()

        # Connect to shared store for cross-talk with other tribunal forms
        self._connect_shared_store()

    def _connect_shared_store(self):
        """Connect to SharedDataStore for cross-report data sharing."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.report_sections_changed.connect(self._on_report_sections_changed)
            shared_store.notes_changed.connect(self._on_notes_changed)
            shared_store.extracted_data_changed.connect(self._on_extracted_data_changed)
            shared_store.patient_info_changed.connect(self._on_patient_info_changed)
            print("[SOCIAL] Connected to SharedDataStore signals (sections, notes, extracted_data, patient_info)")

            # Check if there's already data in the store
            self._check_shared_store_for_existing_data()
        except Exception as e:
            print(f"[SOCIAL] Failed to connect to SharedDataStore: {e}")

    def _check_shared_store_for_existing_data(self):
        """Check SharedDataStore for existing data when page is created."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()

            # Check for existing report sections (cross-talk)
            existing_sections = shared_store.report_sections
            source = shared_store.get_report_source()
            if existing_sections and source != "social_tribunal":
                print(f"[SOCIAL] Found existing sections from {source}, populating...")
                self._on_report_sections_changed(existing_sections, source)

            # Check for existing notes (only if no report data)
            if not self._has_report_data():
                notes = shared_store.notes
                if notes:
                    print(f"[SOCIAL] Found {len(notes)} existing notes in SharedDataStore")
                    if not hasattr(self, '_extracted_raw_notes'):
                        self._extracted_raw_notes = []
                    self._extracted_raw_notes = notes

                # Check for existing extracted data
                extracted_data = shared_store.extracted_data
                if extracted_data:
                    print(f"[SOCIAL] Found existing extracted data in SharedDataStore")
                    self._on_extracted_data_changed(extracted_data)
        except Exception as e:
            print(f"[SOCIAL] Error checking shared store: {e}")

    def _on_patient_info_changed(self, patient_info: dict):
        """Handle patient info updates from SharedDataStore."""
        if patient_info and any(patient_info.values()):
            print(f"[SOCIAL] Received patient info from SharedDataStore: {list(k for k,v in patient_info.items() if v)}")
            self._fill_patient_details(patient_info)

    def _has_report_data(self):
        """Check if report data has been imported (local or via SharedDataStore)."""
        if hasattr(self, '_imported_report_data') and self._imported_report_data:
            return True
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            source = shared_store.get_report_source()
            if source and source != "social_tribunal" and shared_store.report_sections:
                return True
        except Exception:
            pass
        return False

    def _on_notes_changed(self, notes: list):
        """Handle notes updates from SharedDataStore."""
        if notes:
            # Skip if report data exists (report takes priority)
            if self._has_report_data():
                print(f"[SOCIAL] Skipping notes from SharedDataStore - report data already imported")
                return
            print(f"[SOCIAL] Received {len(notes)} notes from SharedDataStore")
            if not hasattr(self, '_extracted_raw_notes'):
                self._extracted_raw_notes = []
            self._extracted_raw_notes = notes

    def _on_extracted_data_changed(self, data: dict):
        """Handle extracted data updates from SharedDataStore."""
        if not data:
            return
        # Skip if report data exists (report takes priority)
        if self._has_report_data():
            print(f"[SOCIAL] Skipping extracted data from SharedDataStore - report data already imported")
            return
        print(f"[SOCIAL] Received extracted data from SharedDataStore: {list(data.keys())}")
        if not hasattr(self, '_extracted_categories'):
            self._extracted_categories = {}
        categories = data.get("categories", data)
        self._extracted_categories = categories

    def _on_report_sections_changed(self, sections: dict, source_form: str):
        """Handle report sections imported from another form (cross-talk)."""
        if source_form == "social_tribunal":
            return

        print(f"[SOCIAL] Cross-talk received from {source_form}: {len(sections)} sections")

        # Store imported data for popups to use
        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        for key, content in sections.items():
            if key in self.cards and content:
                # Store the imported data for the popup's imported data section
                self._imported_report_data[key] = content

                # If popup already exists, populate it with report data
                if hasattr(self, 'popups') and key in self.popups:
                    self._populate_single_popup(self.popups[key], key, content)
                    print(f"[SOCIAL] Populated popup '{key}' from cross-talk")

        print(f"[SOCIAL] Cross-talk stored {len(sections)} sections from {source_form} (cards not auto-filled)")

    def _check_shared_store_for_sections(self):
        """Check SharedDataStore for existing sections when form is shown."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            existing_sections = shared_store.report_sections
            source = shared_store.get_report_source()
            if existing_sections and source and source != "social_tribunal":
                print(f"[SOCIAL] showEvent: Found sections from {source}")
                self._on_report_sections_changed(existing_sections, source)
        except Exception as e:
            print(f"[SOCIAL] Error checking shared store: {e}")

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {
            "full_name": details[1] or "",
            "role_title": details[2] or "",
            "discipline": details[3] or "",
            "registration_body": details[4] or "",
            "registration_number": details[5] or "",
            "phone": details[6] or "",
            "email": details[7] or "",
            "team_service": details[8] or "",
            "hospital_org": details[9] or "",
            "ward_department": details[10] or "",
            "signature_block": details[11] or "",
        }

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar (amber/orange for social)
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("""
            QFrame {
                background: #b45309;
                border-bottom: 1px solid #92400e;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: white;
                border: 1px solid rgba(255,255,255,0.3);
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 500;
            }
            QPushButton:hover { background: rgba(255,255,255,0.1); }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Social Circumstances Tribunal Report")
        title.setStyleSheet("font-size: 20px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Report - Start New")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 16px;
            }
            QPushButton:hover { background: #7f1d1d; }
        """)
        clear_btn.clicked.connect(self._clear_report)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = SocialToolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        # Connect uploaded docs menu to SharedDataStore
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())
        # Helper to get active editor (persists when toolbar clicked)
        def get_active_editor():
            return self._active_editor

        # Helper to safely call editor method
        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        # Connect formatting signals
        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(
                        self,
                        "Spell Check",
                        "No spelling errors found."
                    )

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("QScrollArea { background: #f3f4f6; border: none; }")
        self.main_splitter.addWidget(self.cards_holder)

        self.editor_root = QWidget()
        self.editor_root.setStyleSheet("background: #f3f4f6;")
        self.editor_layout = QVBoxLayout(self.editor_root)
        self.editor_layout.setContentsMargins(32, 24, 32, 24)
        self.editor_layout.setSpacing(16)
        self.cards_holder.setWidget(self.editor_root)

        # Right: Panel (resizable)
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(350)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame {
                background: rgba(245,245,245,0.98);
                border-left: 1px solid rgba(0,0,0,0.08);
            }
        """)
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 450])

        # Connect splitter movement to update card widths
        self.main_splitter.splitterMoved.connect(self._on_splitter_moved)

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        # Header row with title and lock button
        self.panel_header = QWidget()
        self.panel_header.setStyleSheet("""
            background: rgba(245, 158, 11, 0.15);
            border-radius: 8px;
        """)
        header_layout = QHBoxLayout(self.panel_header)
        header_layout.setContentsMargins(12, 8, 12, 8)
        header_layout.setSpacing(8)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 22px;
            font-weight: 700;
            color: #b45309;
            background: transparent;
        """)
        header_layout.addWidget(self.panel_title, 1)

        # Lock button in header
        self.header_lock_btn = QPushButton("Unlocked")
        self.header_lock_btn.setFixedSize(70, 26)
        self.header_lock_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.header_lock_btn.setToolTip("Click to lock this section")
        self.header_lock_btn.setStyleSheet("""
            QPushButton {
                background: rgba(34, 197, 94, 0.3);
                border: 2px solid #22c55e;
                border-radius: 13px;
                font-size: 13px;
                font-weight: 600;
                color: #16a34a;
            }
            QPushButton:hover { background: rgba(34, 197, 94, 0.5); }
        """)
        self.header_lock_btn.clicked.connect(self._toggle_current_popup_lock)
        self.header_lock_btn.hide()
        header_layout.addWidget(self.header_lock_btn)

        panel_layout.addWidget(self.panel_header)

        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        self._create_cards()

    # Sections that should be headings (no editor card, just clickable title)
    HEADING_ONLY_SECTIONS = {
        "s2_detention",      # Section 25
        "other_detention",   # Section 26
    }

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _toggle_current_popup_lock(self):
        """Toggle lock on the currently active popup."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'toggle_lock'):
            popup.toggle_lock()
            self._update_header_lock_button()

    def _update_header_lock_button(self):
        """Update header lock button to match current popup state."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'is_locked'):
            is_locked = popup.is_locked()
            if is_locked:
                self.header_lock_btn.setText("Locked")
                self.header_lock_btn.setToolTip("Click to unlock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(239, 68, 68, 0.3);
                        border: 2px solid #ef4444;
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: #dc2626;
                    }
                    QPushButton:hover { background: rgba(239, 68, 68, 0.5); }
                """)
            else:
                self.header_lock_btn.setText("Unlocked")
                self.header_lock_btn.setToolTip("Click to lock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(34, 197, 94, 0.3);
                        border: 2px solid #22c55e;
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: #16a34a;
                    }
                    QPushButton:hover { background: rgba(34, 197, 94, 0.5); }
                """)
            self.header_lock_btn.show()
        else:
            self.header_lock_btn.hide()

    def _set_current_popup(self, popup):
        """Set the current active popup and update lock button."""
        self._current_popup = popup
        self._update_header_lock_button()

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self  # Capture reference to self for closure

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _create_cards(self):
        """Create all section cards (or headings for certain sections)."""
        for title, key in self.SECTIONS:
            if key in self.HEADING_ONLY_SECTIONS:
                # Create heading widget instead of full card
                card = SocialHeadingWidget(title, key, parent=self.editor_root)
            else:
                # Create standard card widget
                card = SocialCardWidget(title, key, parent=self.editor_root)
                # Hook up focus event to register this editor as active
                self._hook_editor_focus(card.editor)
                # Connect card text changes to sync with popup
                card.editor.textChanged.connect(lambda k=key: self._on_card_text_changed(k))
            card.clicked.connect(self._on_card_clicked)
            self.cards[key] = card
            self.editor_layout.addWidget(card)
        self.editor_layout.addStretch()

    def _on_splitter_moved(self, pos, index):
        """Update card widths when splitter is moved."""
        self._update_card_widths()

    def _update_card_widths(self):
        """Update card widths based on current viewport size."""
        viewport_width = self.cards_holder.viewport().width()
        # Account for margins (24px on each side)
        card_width = viewport_width - 48
        if card_width > 100:  # Minimum sensible width
            for key, card in self.cards.items():
                if key in self.HEADING_ONLY_SECTIONS:
                    # Headings use max width so text wraps
                    card.setMaximumWidth(card_width)
                else:
                    card.setFixedWidth(card_width)

    def showEvent(self, event):
        """Set initial card widths when page is shown."""
        super().showEvent(event)
        # Use a timer to ensure layout is complete
        from PySide6.QtCore import QTimer
        QTimer.singleShot(50, self._update_card_widths)

        # Check for existing report sections in shared store (cross-talk)
        QTimer.singleShot(100, self._check_shared_store_for_sections)

    def resizeEvent(self, event):
        """Update card widths on window resize."""
        super().resizeEvent(event)
        self._update_card_widths()

    def _on_card_clicked(self, key: str):
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        print(f"[SOCIAL] Card clicked: {key}")

        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)
        if key in self.cards:
            self.cards[key].setSelected(True)
        self._selected_card_key = key

        if key not in self.popups:
            popup = self._create_popup(key)
            if popup:
                self.popups[key] = popup
                self.popup_stack.addWidget(popup)

                if key == "forensic":
                    popup.sent.connect(lambda text, state, k=key: self._update_card(k, text))
                elif key in ("patient_details", "home_family", "housing", "financial", "employment", "signature"):
                    # These popups send on every change - use replace mode
                    popup.sent.connect(lambda text, k=key: self._set_card(k, text))
                    # Connect gender_changed for patient_details
                    if key == "patient_details" and hasattr(popup, 'gender_changed'):
                        popup.gender_changed.connect(self._on_gender_changed)
                else:
                    popup.sent.connect(lambda text, k=key: self._update_card(k, text))

                # Populate from imported report data if available
                if hasattr(self, '_imported_report_data') and key in self._imported_report_data:
                    if not getattr(popup, '_imported_data_added', False):
                        content = self._imported_report_data[key]
                        self._populate_single_popup(popup, key, content)
                        print(f"[SOCIAL] Populated popup '{key}' from imported report data")

                # Notes-based searches - only run when NO report data exists
                if not self._has_report_data():
                    # Check for pending section data first
                    if hasattr(self, '_pending_section_data') and key in self._pending_section_data:
                        entries = self._pending_section_data[key]
                        if key in ["risk_harm", "risk_property"] and hasattr(popup, 'set_notes'):
                            popup.set_notes(entries)
                            print(f"[SOCIAL] Loaded pending notes for '{key}' ({len(entries)} notes)")
                        elif hasattr(popup, 'set_entries'):
                            popup.set_entries(entries, f"{len(entries)} entries")
                            print(f"[SOCIAL] Loaded pending data for '{key}' ({len(entries)} entries)")
                    elif key in ["risk_harm", "risk_property"]:
                        notes = getattr(self, '_extracted_raw_notes', [])
                        if hasattr(popup, 'set_notes') and notes:
                            popup.set_notes(notes)
                            print(f"[SOCIAL] Loaded raw notes for '{key}' ({len(notes)} notes)")

                    # For forensic popup - populate with forensic data
                    if key == "forensic":
                        forensic_entries = []
                        if hasattr(self, '_pending_forensic_data') and 'forensic' in self._pending_forensic_data:
                            forensic_entries = self._pending_forensic_data['forensic']
                        notes = getattr(self, '_extracted_raw_notes', [])
                        if hasattr(popup, 'set_forensic_data'):
                            popup.set_forensic_data(notes, forensic_entries)
                            print(f"[SOCIAL] Populated forensic panel with {len(notes)} notes and {len(forensic_entries)} entries")

                    # For previous_mh_dates - run timeline analysis
                    if key == "previous_mh_dates":
                        notes = getattr(self, '_extracted_raw_notes', [])
                        if hasattr(popup, 'set_notes') and notes:
                            popup.set_notes(notes)
                            print(f"[SOCIAL] Ran timeline analysis for section 5 ({len(notes)} notes)")

                    # For discharge_risk - risk analysis
                    if key == "discharge_risk":
                        notes = getattr(self, '_extracted_raw_notes', [])
                        if notes and hasattr(popup, 'set_notes_for_risk_analysis'):
                            popup.set_notes_for_risk_analysis(notes)
                            print(f"[SOCIAL] Populated discharge_risk with {len(notes)} notes for risk analysis")

        if key in self.popups:
            self.popup_stack.setCurrentWidget(self.popups[key])
            self._set_current_popup(self.popups[key])

            # Send popup form content to card (imported data checkbox is unchecked so only form data flows)
            if hasattr(self, '_imported_report_data') and key in self._imported_report_data:
                popup = self.popups[key]
                if hasattr(popup, '_send_to_card'):
                    popup._send_to_card()

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details popup and card from extracted demographics."""
        if not patient_info:
            return

        # Create popup if it doesn't exist
        if "patient_details" not in self.popups:
            from tribunal_popups import PatientDetailsPopup
            popup = PatientDetailsPopup(parent=self)
            self.popups["patient_details"] = popup
            self.popup_stack.addWidget(popup)
            # Connect sent signal to replace card content
            popup.sent.connect(lambda text: self._set_card("patient_details", text))
            # Connect gender_changed to update all popups
            popup.gender_changed.connect(self._on_gender_changed)

        # Fill the popup fields
        popup = self.popups["patient_details"]
        if hasattr(popup, 'fill_patient_info'):
            popup.fill_patient_info(patient_info)

            # Also update the card preview
            text = popup.generate_text()
            if "patient_details" in self.cards and text.strip():
                self.cards["patient_details"].editor.setPlainText(text)
                print(f"[SocialTribunalReport] Updated patient_details card with demographics")

    def _create_popup(self, key: str):
        from tribunal_popups import PatientDetailsPopup, FactorsHearingPopup, AdjustmentsPopup
        from tribunal_popups import StrengthsPopup, DoLsPopup, YesNoNAPopup, SimpleYesNoPopup
        from tribunal_popups import DischargeRiskPopup, CommunityManagementPopup
        from tribunal_popups import RecommendationsPopup, SignaturePopup
        from forensic_history_popup import ForensicHistoryPopup
        from tribunal_popups import TribunalPsychHistoryPopup

        if key == "patient_details":
            return PatientDetailsPopup(parent=self)
        elif key == "factors_hearing":
            return FactorsHearingPopup(parent=self, gender=self._current_gender)
        elif key == "adjustments":
            return AdjustmentsPopup(parent=self, gender=self._current_gender)
        elif key == "forensic":
            return ForensicHistoryPopup(parent=self, gender=self._current_gender, show_index_offence=True)
        elif key == "previous_mh_dates":
            return TribunalPsychHistoryPopup(parent=self)
        elif key == "home_family":
            return HomeFamilyPopup(parent=self, gender=self._current_gender)
        elif key == "housing":
            return HousingPopup(parent=self, gender=self._current_gender)
        elif key == "financial":
            return FinancialPopup(parent=self, gender=self._current_gender)
        elif key == "employment":
            return EmploymentPopup(parent=self, gender=self._current_gender)
        elif key == "previous_community":
            return PreviousCommunityPopup(parent=self, gender=self._current_gender)
        elif key == "care_pathway":
            return CarePathwayPopup(parent=self, gender=self._current_gender)
        elif key == "care_plan":
            return CarePlanPopup(parent=self, gender=self._current_gender)
        elif key == "care_plan_adequacy":
            return CarePlanAdequacyPopup(parent=self)
        elif key == "care_plan_funding":
            return CarePlanFundingPopup(parent=self)
        elif key == "strengths":
            return StrengthsPopup(parent=self, gender=self._current_gender)
        elif key == "progress":
            # Section 16: Use TribunalProgressPopup identical to psych tribunal section 14
            # Filter to 1 year from most recent entry
            from tribunal_popups import TribunalProgressPopup
            return TribunalProgressPopup(parent=self, date_filter='1_year')
        elif key == "risk_harm":
            # Section 17: Use TribunalRiskHarmPopup identical to psych tribunal section 17
            from tribunal_popups import TribunalRiskHarmPopup
            return TribunalRiskHarmPopup(parent=self)
        elif key == "risk_property":
            # Section 18: Use TribunalRiskPropertyPopup identical to psych tribunal section 18
            from tribunal_popups import TribunalRiskPropertyPopup
            return TribunalRiskPropertyPopup(parent=self)
        elif key == "patient_views":
            return PatientViewsPopup(parent=self, gender=self._current_gender)
        elif key == "nearest_relative":
            return NearestRelativePopup(parent=self, gender=self._current_gender)
        elif key == "nr_inappropriate":
            return NRInappropriatePopup(parent=self, gender=self._current_gender)
        elif key == "carer_views":
            return CarerViewsPopup(parent=self, gender=self._current_gender)
        elif key == "mappa":
            return MAPPAPopup(parent=self, gender=self._current_gender)
        elif key == "mca_dol":
            return DoLsPopup(parent=self)
        elif key == "s2_detention":
            return YesNoNAPopup(
                title="Section 2 Detention",
                question="Is detention under Section 2 justified for health, safety or protection of others?",
                yes_output="Yes, detention under Section 2 is necessary.",
                no_output="No, detention under Section 2 is not necessary.",
                na_output="Not applicable - patient is not detained under Section 2.",
                parent=self
            )
        elif key == "other_detention":
            return SimpleYesNoPopup(
                title="Medical Treatment Justified",
                question="Is the provision of medical treatment in hospital justified or necessary in the interests of the patient's health or safety, or for the protection of others?",
                parent=self
            )
        elif key == "discharge_risk":
            # Section 27: Use GPRRiskPopup identical to psych tribunal section 21
            from general_psychiatric_report_page import GPRRiskPopup
            return GPRRiskPopup(parent=self)
        elif key == "community":
            return CommunityManagementPopup(parent=self, gender=self._current_gender)
        elif key == "other_info":
            return OtherInfoPopup(parent=self)
        elif key == "recommendations":
            try:
                from general_psychiatric_report_page import GPRLegalCriteriaPopup
                from icd10_dict import ICD10_DICT
                print(f"[SOCIAL] Creating GPRLegalCriteriaPopup: ICD10_DICT={len(ICD10_DICT)} entries, gender={self._current_gender}")
                popup = GPRLegalCriteriaPopup(parent=self, gender=self._current_gender, icd10_dict=ICD10_DICT)
                print(f"[SOCIAL] GPRLegalCriteriaPopup created successfully")
                return popup
            except Exception as e:
                print(f"[SOCIAL] ERROR creating recommendations popup: {e}")
                import traceback
                traceback.print_exc()
                from PySide6.QtWidgets import QLabel
                err = QLabel(f"ERROR loading recommendations popup:\n{e}")
                err.setStyleSheet("color: red; font-size: 16px; padding: 20px;")
                err.setWordWrap(True)
                return err
        elif key == "signature":
            return SignaturePopup(parent=self, my_details=self._my_details)

        return None

    def _update_card(self, key: str, text: str):
        """Replace card content with new text from popup."""
        if key in self.cards:
            self.cards[key].editor.setPlainText(text.strip() if text else "")

    def _set_card(self, key: str, text: str):
        """Replace card content (used by popups that send on every checkbox change)."""
        if key in self.cards:
            self.cards[key].editor.setPlainText(text.strip() if text else "")

    def _on_card_text_changed(self, key: str):
        """Handle card text changes - sync to popup if not already syncing."""
        if getattr(self, '_syncing', False):
            return

        if key not in self.cards:
            return

        card_text = self.cards[key].editor.toPlainText()

        # Parse card text and update corresponding popup
        if key in self.popups:
            self._syncing = True
            try:
                self._update_popup_from_card(key, card_text)
            finally:
                self._syncing = False

    def _update_popup_from_card(self, key: str, card_text: str):
        """Update popup fields from card text."""
        if key not in self.popups:
            return

        popup = self.popups[key]
        fields = self._parse_structured_text(card_text)

        # Patient details popup
        if key == "patient_details":
            # Handle name field
            name_value = fields.get('Full Name') or fields.get('Name') or ''
            if name_value and hasattr(popup, 'name_field'):
                popup.name_field.blockSignals(True)
                popup.name_field.setText(name_value)
                popup.name_field.blockSignals(False)

            # Handle DOB field
            dob_value = fields.get('Date of Birth') or fields.get('DOB') or ''
            if dob_value and hasattr(popup, 'dob_field'):
                popup.dob_field.blockSignals(True)
                from PySide6.QtCore import QDate
                for fmt in ["dd/MM/yyyy", "dd-MM-yyyy", "yyyy-MM-dd", "d MMMM yyyy"]:
                    parsed = QDate.fromString(dob_value.strip(), fmt)
                    if parsed.isValid():
                        popup.dob_field.setDate(parsed)
                        break
                popup.dob_field.blockSignals(False)

            # Handle Gender field
            gender_value = fields.get('Gender') or ''
            if gender_value and hasattr(popup, 'gender_field'):
                popup.gender_field.blockSignals(True)
                idx = popup.gender_field.findText(gender_value, Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    popup.gender_field.setCurrentIndex(idx)
                popup.gender_field.blockSignals(False)
                # Propagate gender since signals were blocked
                self._on_gender_changed(gender_value)

            # Handle Residence field
            residence_value = fields.get('Usual Place of Residence') or fields.get('Residence') or fields.get('Address') or ''
            if residence_value and hasattr(popup, 'residence_field'):
                popup.residence_field.blockSignals(True)
                popup.residence_field.setPlainText(residence_value)
                popup.residence_field.blockSignals(False)

    def _parse_structured_text(self, text: str) -> dict:
        """Parse text with 'Field: Value' format into a dictionary."""
        fields = {}
        for line in text.split('\n'):
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    field_name = parts[0].strip()
                    value = parts[1].strip()
                    if field_name and value:
                        fields[field_name] = value
        return fields

    def _on_gender_changed(self, gender: str):
        """Update gender and propagate to all existing popups."""
        self._current_gender = gender
        print(f"[SOCIAL] Gender changed to: {gender}")

        # Update all existing popups that support gender
        for key, popup in self.popups.items():
            if hasattr(popup, 'update_gender'):
                popup.update_gender(gender)
            elif hasattr(popup, 'set_gender'):
                popup.set_gender(gender)
                # Trigger refresh if popup has _send_to_card
                if hasattr(popup, '_send_to_card'):
                    popup._send_to_card()

    def _go_back(self):
        self.go_back.emit()

    def _clear_report(self):
        reply = QMessageBox.question(
            self, "Clear Report",
            "Are you sure you want to clear all data and start a new report?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            for card in self.cards.values():
                if hasattr(card, 'editor'):
                    card.editor.clear()
            self._extracted_raw_notes = []
            self._extracted_categories = {}
            self._incident_data = []
            self._data_processed_id = None
            self._notes_processed_id = None
            if hasattr(self, '_imported_report_data'):
                self._imported_report_data = {}
            if hasattr(self, '_imported_report_sections'):
                self._imported_report_sections = {}

            # Destroy all popups and remove from stack
            for key, popup in list(self.popups.items()):
                if hasattr(self, 'popup_stack'):
                    self.popup_stack.removeWidget(popup)
                popup.deleteLater()
            self.popups.clear()

            # Clear data extractor
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                if hasattr(self._data_extractor_overlay, 'clear_extraction'):
                    self._data_extractor_overlay.clear_extraction()

            # Recreate signature popup so mydetails get restored
            if hasattr(self, '_on_card_clicked'):
                if "signature" in self.cards:
                    self._on_card_clicked("signature")

            print("[SOCIAL] Report cleared - ready for new report")

    def _export_docx(self):
        """Export report to DOCX using T133 template with table-based input boxes."""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from docx import Document
        from docx.shared import Pt
        from datetime import datetime
        import os
        import shutil
        import re

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Social Circumstances Tribunal Report",
            f"social_tribunal_report_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            # Find template path
            template_path = resource_path("templates", "t133_template_new.docx")

            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Template Error", f"T133 template not found at:\n{template_path}")
                return

            # Copy template to destination
            shutil.copy(template_path, file_path)

            # Open the copied template
            doc = Document(file_path)

            # Helper to get card content (or popup state for heading-only sections)
            def get_content(key):
                # For heading-only sections, get content from popup's generate_text()
                if key in self.HEADING_ONLY_SECTIONS:
                    if key in self.popups and hasattr(self.popups[key], 'generate_text'):
                        return self.popups[key].generate_text()
                    return ""
                # For regular cards, get from editor
                if key in self.cards:
                    return self.cards[key].editor.toPlainText().strip()
                return ""

            # Helper to extract detail text after Yes/No prefix
            def extract_detail(content):
                if not content:
                    return ""
                detail = content
                lower = detail.lower()
                for prefix in ["yes, ", "yes - ", "yes-", "yes,", "yes.", "no, ", "no - ", "no-", "no,", "no."]:
                    if lower.startswith(prefix):
                        detail = detail[len(prefix):].strip()
                        break
                if lower == "yes" or lower == "no":
                    return ""
                return detail

            # Helper to check if content indicates Yes answer
            def has_yes_content(content):
                if not content:
                    return False
                lower = content.lower().strip()
                no_patterns = [
                    "no ", "no,", "no-", "no.",
                    "there have been no", "there is no",
                    "no incidents", "no factors", "no adjustments",
                    "does not have", "not applicable", "n/a",
                ]
                if lower == "no":
                    return False
                for pattern in no_patterns:
                    if lower.startswith(pattern) or pattern in lower:
                        return False
                return True

            # Helper to set table cell text
            def set_cell(table_idx, row, col, text):
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    if row < len(table.rows) and col < len(table.rows[row].cells):
                        cell = table.cell(row, col)
                        cell.text = text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(11)

            # ============================================================
            # PARSE PATIENT DETAILS
            # ============================================================
            pd_content = get_content("patient_details")
            pd_lines = pd_content.split('\n') if pd_content else []
            full_name = dob = residence = ""

            for line in pd_lines:
                line_lower = line.lower()
                if "name:" in line_lower or "full name:" in line_lower:
                    full_name = line.split(":", 1)[-1].strip()
                elif "date of birth:" in line_lower or "dob:" in line_lower:
                    dob = line.split(":", 1)[-1].strip()
                elif any(x in line_lower for x in ["residence:", "address:", "usual place"]):
                    residence = line.split(":", 1)[-1].strip()

            # If no structured name, use first non-empty line
            if not full_name and pd_lines:
                for line in pd_lines:
                    if line.strip() and not any(x in line.lower() for x in ["dob", "date of birth", "residence", "address"]):
                        full_name = line.strip()
                        break

            # Table 1: Full name (1x2) - name goes in second column
            set_cell(1, 0, 1, full_name)

            # Table 2: DOB character boxes (1x8)
            if dob:
                dob_clean = ""
                date_formats = [
                    ("%d/%m/%Y", None), ("%d-%m-%Y", None), ("%d.%m.%Y", None),
                    ("%d %B %Y", None), ("%d %b %Y", None), ("%Y-%m-%d", None),
                ]
                parsed_date = None
                for fmt, _ in date_formats:
                    try:
                        parsed_date = datetime.strptime(dob.strip(), fmt)
                        break
                    except:
                        continue

                if parsed_date:
                    dob_clean = parsed_date.strftime("%d%m%Y")
                else:
                    dob_clean = re.sub(r'[^0-9]', '', dob)
                    if len(dob_clean) == 6:
                        year_part = dob_clean[4:6]
                        year = "20" + year_part if int(year_part) < 50 else "19" + year_part
                        dob_clean = dob_clean[:4] + year

                if len(doc.tables) > 2:
                    dob_table = doc.tables[2]
                    for i, char in enumerate(dob_clean[:8]):
                        if i < len(dob_table.columns):
                            dob_table.cell(0, i).text = char

            # Table 3: Residence (1x1)
            set_cell(3, 0, 0, residence)

            # ============================================================
            # PARSE SIGNATURE INFO
            # ============================================================
            sig_content = get_content("signature")
            sig_lines = sig_content.split('\n') if sig_content else []
            sig_name = sig_role = ""
            sig_date = datetime.now().strftime("%d/%m/%Y")

            for line in sig_lines:
                line_lower = line.lower()
                if "signed:" in line_lower or "name:" in line_lower:
                    sig_name = line.split(":", 1)[-1].strip()
                elif "designation:" in line_lower or "role:" in line_lower:
                    sig_role = line.split(":", 1)[-1].strip()
                elif "date:" in line_lower:
                    sig_date = line.split(":", 1)[-1].strip()

            if not sig_name and sig_lines:
                for line in sig_lines:
                    if line.strip() and not any(x in line.lower() for x in ["date", "designation", "role", "registration"]):
                        sig_name = line.strip()
                        break

            # Get author name if available
            author_name = get_content("author") or sig_name

            # ============================================================
            # NESTED TABLE IN TABLE 0 - Hospital info box (top right)
            # Structure: Row 0=Hospital logo title, Row 1=Your name title,
            #            Row 2=name entry, Row 3=Your role title,
            #            Row 4=role entry, Row 5=Date title, Row 6=date entry
            # ============================================================
            if len(doc.tables) > 0:
                table0 = doc.tables[0]
                if len(table0.rows[0].cells) > 1:
                    cell_with_nested = table0.rows[0].cells[1]
                    if cell_with_nested.tables:
                        nested_table = cell_with_nested.tables[0]
                        # Row 2: Your name entry
                        if len(nested_table.rows) > 2:
                            nested_table.rows[2].cells[0].text = author_name
                        # Row 4: Your role entry
                        if len(nested_table.rows) > 4:
                            nested_table.rows[4].cells[0].text = sig_role
                        # Row 6: Date entry
                        if len(nested_table.rows) > 6:
                            nested_table.rows[6].cells[0].text = sig_date

            # ============================================================
            # SECTION INPUT BOXES
            # T133 Table mapping: table_idx -> (card_key, use_extract_detail)
            # ============================================================
            table_mapping = {
                4: ("factors_hearing", True),
                5: ("adjustments", True),
                6: ("forensic", False),
                7: ("previous_mh_dates", False),
                8: ("home_family", False),
                9: ("housing", False),
                10: ("financial", False),
                11: ("employment", True),
                12: ("previous_community", False),
                13: ("care_pathway", False),
                14: ("care_plan", False),
                15: ("care_plan_adequacy", False),
                # Table 16 is funding date boxes (1x8) - handled separately
                17: ("strengths", False),
                18: ("progress", False),
                19: ("risk_harm", False),
                20: ("risk_property", False),
                21: ("patient_views", False),
                22: ("nearest_relative", False),
                23: ("nr_inappropriate", False),
                24: ("carer_views", False),
                25: ("mappa", True),
                # Table 26: MAPPA chair name - handled separately
                # Table 27: MAPPA lead agency - handled separately
                28: ("mca_dol", False),
                29: ("discharge_risk", True),  # Combined with community
                30: ("other_info", True),
                31: ("recommendations", True),
                # Table 32: Signature
                # Table 33: Date boxes
            }

            # Sections with Yes/No checkboxes - only fill table if Yes
            yes_no_sections = {"factors_hearing", "adjustments", "employment",
                               "mappa", "discharge_risk", "other_info", "recommendations"}

            for table_idx, (card_key, use_extract) in table_mapping.items():
                content = get_content(card_key)
                if content:
                    if card_key in yes_no_sections:
                        if not has_yes_content(content):
                            continue
                    if use_extract:
                        content = extract_detail(content) or content
                    set_cell(table_idx, 0, 0, content)

            # Handle community content - append to discharge_risk table if present
            community_content = get_content("community")
            if community_content:
                existing = doc.tables[29].cell(0, 0).text if len(doc.tables) > 29 else ""
                if existing:
                    set_cell(29, 0, 0, existing + "\n\n" + community_content)
                else:
                    set_cell(29, 0, 0, community_content)

            # Table 16: Funding date boxes (if care_plan_funding has a date)
            funding_content = get_content("care_plan_funding")
            if funding_content and has_yes_content(funding_content):
                # Try to extract a date from the funding content
                date_match = re.search(r'(\d{1,2})[/.-](\d{1,2})[/.-](\d{2,4})', funding_content)
                if date_match and len(doc.tables) > 16:
                    day, month, year = date_match.groups()
                    if len(year) == 2:
                        year = "20" + year if int(year) < 50 else "19" + year
                    date_str = f"{day.zfill(2)}{month.zfill(2)}{year}"
                    funding_table = doc.tables[16]
                    for i, char in enumerate(date_str[:8]):
                        if i < len(funding_table.columns):
                            funding_table.cell(0, i).text = char

            # ============================================================
            # CHECKBOX SECTIONS - mark ☐ → ☒ based on content
            # ============================================================
            checkbox_patterns = {
                "factors that may affect": "factors_hearing",
                "any adjustments": "adjustments",
                "opportunities for employment": "employment",
                "issues as to funding": "care_plan_funding",
                "mappa meeting": "mappa",
                "section 2 cases": "s2_detention",
                "in all other cases": "other_detention",
                "discharged from hospital": "discharge_risk",
                "other relevant information": "other_info",
                "recommendations to the tribunal": "recommendations",
            }

            current_section_key = None
            for para in doc.paragraphs:
                text = para.text.strip()
                lower_text = text.lower()

                for pattern, card_key in checkbox_patterns.items():
                    if pattern in lower_text:
                        current_section_key = card_key
                        break

                if current_section_key and "☐" in para.text:
                    content = get_content(current_section_key)
                    is_yes = has_yes_content(content)

                    is_yes_line = "yes" in lower_text and "no" not in lower_text[:10]
                    is_no_line = lower_text.strip().startswith("☐ no") or (lower_text.strip().startswith("☐") and "no" in lower_text and "yes" not in lower_text)

                    if is_yes_line and is_yes:
                        for run in para.runs:
                            if "☐" in run.text:
                                run.text = run.text.replace("☐", "☒", 1)
                                break
                    elif is_no_line and not is_yes and content:
                        for run in para.runs:
                            if "☐" in run.text:
                                run.text = run.text.replace("☐", "☒", 1)
                                break

            # ============================================================
            # SIGNATURE - Table 32
            # ============================================================
            from docx.shared import Inches

            sig_text_parts = []
            if sig_name:
                sig_text_parts.append(sig_name)
            if sig_role:
                sig_text_parts.append(sig_role)

            for line in sig_lines:
                line_lower = line.lower()
                if "qualification" in line_lower:
                    sig_text_parts.append(line.split(":", 1)[-1].strip())
                elif "registration:" in line_lower or "hcpc:" in line_lower or "sweng:" in line_lower:
                    reg = line.split(":", 1)[-1].strip()
                    sig_text_parts.append(f"Registration: {reg}")

            # Put signature image and text in Table 32
            if 32 < len(doc.tables):
                sig_cell = doc.tables[32].cell(0, 0)
                sig_cell.text = ""  # Clear existing content

                # Add signature image if exists
                sig_image_path = os.path.expanduser("~/MyPsychAdmin/signature.png")
                if os.path.exists(sig_image_path):
                    sig_para = sig_cell.paragraphs[0]
                    sig_run = sig_para.add_run()
                    sig_run.add_picture(sig_image_path, width=Inches(1.5))
                    # Add new paragraph for text
                    sig_cell.add_paragraph("\n".join(sig_text_parts))
                else:
                    sig_cell.text = "\n".join(sig_text_parts)

            # ============================================================
            # TABLE 33: Date character boxes (1x8) - ddmmyyyy format
            # ============================================================
            if sig_date and len(doc.tables) > 33:
                date_table = doc.tables[33]
                date_clean = ""

                date_formats = [
                    ("%d/%m/%Y", None), ("%d-%m-%Y", None), ("%d.%m.%Y", None),
                    ("%d %B %Y", None), ("%d %b %Y", None), ("%Y-%m-%d", None),
                    ("%B %d, %Y", None), ("%b %d, %Y", None),
                ]

                parsed_date = None
                for fmt, _ in date_formats:
                    try:
                        parsed_date = datetime.strptime(sig_date.strip(), fmt)
                        break
                    except:
                        continue

                if parsed_date:
                    date_clean = parsed_date.strftime("%d%m%Y")
                else:
                    date_clean = re.sub(r'[^0-9]', '', sig_date)
                    if len(date_clean) == 6:
                        year_part = date_clean[4:6]
                        year = "20" + year_part if int(year_part) < 50 else "19" + year_part
                        date_clean = date_clean[:4] + year

                for i, char in enumerate(date_clean[:8]):
                    if i < len(date_table.columns):
                        date_table.cell(0, i).text = char

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Report exported to:\n{file_path}")

        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Export Error", f"Failed to export report:\n{str(e)}")

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            from shared_data_store import get_shared_store
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file - auto-detect PDF/DOCX tribunal forms vs other files."""
        # Detect cross-report imports (T131 PTR / T134 NTR) and route to data extractor
        # which uses PTR_TO_SCT_MAP for correct section mapping
        is_cross_report = False
        if file_path.lower().endswith(('.pdf', '.docx')):
            try:
                full_text = ""
                if file_path.lower().endswith('.pdf'):
                    import fitz
                    doc = fitz.open(file_path)
                    full_text = "\n".join(page.get_text() for page in doc)
                    doc.close()
                elif file_path.lower().endswith('.docx'):
                    from docx import Document
                    doc = Document(file_path)
                    full_text = "\n".join(p.text for p in doc.paragraphs)
                lower = full_text.lower()
                if ('responsible clinician' in lower or 't131' in lower or
                        'nursing report' in lower or 't134' in lower):
                    if 'social circumstances' not in lower and 'social supervisor' not in lower:
                        is_cross_report = True
                        print(f"[SOCIAL] Cross-report import detected (T131/T134) - routing to data extractor")
            except Exception as e:
                print(f"[SOCIAL] Report type detection error: {e}")

        if is_cross_report:
            self._send_to_data_extractor(file_path)
            return

        # Check if this is a PDF - might be a social tribunal form
        if file_path.lower().endswith('.pdf'):
            try:
                from pdf_loader import load_tribunal_pdf, format_radio_value

                result = load_tribunal_pdf(file_path)

                # Check if it's a recognized tribunal form with data
                if result.get('sections'):
                    self._populate_from_pdf(result, file_path)
                    return
                else:
                    # Not a tribunal form or empty - send to data extractor
                    self._send_to_data_extractor(file_path)
                    return

            except Exception as e:
                print(f"[SOCIAL] PDF load error: {e}")
                self._send_to_data_extractor(file_path)
                return

        # Check if this is a DOCX - might be a tribunal report
        elif file_path.lower().endswith('.docx'):
            try:
                result = self._parse_tribunal_docx(file_path)
                if result and result.get('sections'):
                    self._populate_from_docx(result, file_path)
                    return
                else:
                    self._send_to_data_extractor(file_path)
                    return
            except Exception as e:
                print(f"[SOCIAL] DOCX load error: {e}")
                self._send_to_data_extractor(file_path)
                return
        else:
            self._send_to_data_extractor(file_path)

    def _get_heading_to_card_mapping(self) -> list:
        """Get heading patterns mapped to social circumstances card keys.

        This enables cross-talk: importing T131 (medical) or T134 (nursing) reports
        and mapping them to T133 (social circumstances) sections by heading content.
        """
        return [
            # NOTE: patient_details handled by special-case row parser, not heading match

            # Factors affecting hearing - all forms have this
            (r'factors.*affect.*understanding', 'factors_hearing'),
            (r'ability.*cope.*hearing', 'factors_hearing'),

            # Adjustments - all forms have this
            (r'adjustments.*panel.*consider', 'adjustments'),
            (r'adjustments.*may.*consider', 'adjustments'),
            (r'adjustments.*tribunal.*consider', 'adjustments'),

            # Forensic/index offence - T131 section 5, T133 section 4
            (r'index offence', 'forensic'),
            (r'relevant.*forensic', 'forensic'),

            # Previous mental health - T131 section 6, T133 section 5
            (r'dates.*previous.*mental health', 'previous_mh_dates'),
            (r'previous.*involvement.*mental health', 'previous_mh_dates'),

            # Home/family circumstances - T133 specific
            (r'home.*family.*circumstances', 'home_family'),
            (r'family.*circumstances', 'home_family'),

            # Housing - T133 specific
            (r'housing.*accommodation', 'housing'),
            (r'accommodation.*available.*patient', 'housing'),

            # Financial - T133 specific
            (r'financial.*position', 'financial'),
            (r'benefit.*entitlement', 'financial'),

            # Employment - T133 specific
            (r'opportunities.*employment', 'employment'),
            (r'employment.*education', 'employment'),
            (r'employment.*opportunities', 'employment'),

            # Previous community support - T133 specific
            (r'previous.*response.*community', 'previous_community'),
            (r'previously.*supported.*community', 'previous_community'),
            (r'community.*support.*previously', 'previous_community'),

            # Care pathway - T133 specific
            (r'care.*pathway.*section 117', 'care_pathway'),
            (r'care.*pathway.*after.?care', 'care_pathway'),
            (r'section 117.*after.?care.*available', 'care_pathway'),
            (r'care.*pathway', 'care_pathway'),

            # Care plan adequacy - T133 specific (MUST be BEFORE care_plan)
            (r'adequate.*effective.*care plan', 'care_plan_adequacy'),
            (r'adequate.*care plan', 'care_plan_adequacy'),
            (r'how adequate', 'care_plan_adequacy'),

            # Care plan details - T133 specific
            (r'details.*proposed.*care plan', 'care_plan'),
            (r'details.*care plan', 'care_plan'),
            (r'proposed care plan', 'care_plan'),

            # Care plan funding - T133 specific
            (r'issues.*funding.*care plan', 'care_plan_funding'),
            (r'funding.*care plan', 'care_plan_funding'),
            (r'funding.*aftercare', 'care_plan_funding'),
            (r'issues.*funding', 'care_plan_funding'),

            # Strengths - all forms have this
            (r'strengths.*positive factors', 'strengths'),
            (r'what are the strengths', 'strengths'),

            # Progress - all forms have this (T131 s14, T134 s9, T133 s16)
            (r'current progress', 'progress'),
            (r'summary.*progress', 'progress'),
            (r'progress.*behaviour', 'progress'),

            # Risk harm - all forms have this
            (r'harmed themselves or others', 'risk_harm'),
            (r'incidents.*harm', 'risk_harm'),

            # Risk property - all forms have this
            (r'damaged property', 'risk_property'),

            # NR inappropriate - T133 (MUST be BEFORE nearest_relative and patient_views)
            (r'inappropriate.*consult.*nearest\s*relative', 'nr_inappropriate'),
            (r'impractical.*consult.*nearest\s*relative', 'nr_inappropriate'),
            (r'inappropriate.*nearest\s*relative', 'nr_inappropriate'),

            # Carer views - T133 (MUST be BEFORE patient_views)
            (r'other person.*lead.*role', 'carer_views'),
            (r'views.*other person', 'carer_views'),
            (r'carer.*views', 'carer_views'),
            (r'views.*carer', 'carer_views'),

            # Nearest relative views - T133 (MUST be BEFORE patient_views)
            (r'views.*nearest\s*relative', 'nearest_relative'),
            (r'nearest\s*relative.*views', 'nearest_relative'),

            # Patient views - T133 specific (AFTER nr_inappropriate, carer_views, nearest_relative)
            (r"patient'?s?\s+views", 'patient_views'),
            (r'views.*wishes.*beliefs', 'patient_views'),

            # MAPPA - T133 specific
            (r'known.*mappa', 'mappa'),
            (r'mappa\s+meeting', 'mappa'),
            (r'multi.*agency.*public.*protection', 'mappa'),

            # MCA/DoL - T131 and T133 have this
            (r'mental capacity act', 'mca_dol'),
            (r'deprivation of liberty', 'mca_dol'),

            # Section 2 detention - all forms
            (r'in section 2 cases', 's2_detention'),
            (r'section 2 cases.*detention', 's2_detention'),

            # Other detention - all forms
            (r'in all other cases', 'other_detention'),
            (r'all other cases.*provision', 'other_detention'),

            # Discharge risk - all forms
            (r'discharged.*dangerous', 'discharge_risk'),
            (r'discharged.*likely to act', 'discharge_risk'),
            (r'likely to act in a manner dangerous', 'discharge_risk'),

            # Recommendations - all forms (MUST be BEFORE community to avoid "managed in community" sub-question)
            (r'recommendations.*tribunal', 'recommendations'),
            (r'do you have any recommendations', 'recommendations'),

            # Community risk management - all forms
            (r'^\d+\.\s*community.*risk', 'community'),

            # Other info - all forms
            (r'other relevant information', 'other_info'),
            (r'other.*information.*tribunal', 'other_info'),
        ]

    def _match_heading_to_card(self, heading_text: str) -> str:
        """Match a heading to a social circumstances card key using regex patterns."""
        import re

        if not heading_text or len(heading_text) < 10:
            return None

        lower = heading_text.lower()
        patterns = self._get_heading_to_card_mapping()

        for pattern, card_key in patterns:
            if re.search(pattern, lower):
                return card_key

        return None

    def _parse_tribunal_docx(self, file_path: str) -> dict:
        """Parse a DOCX tribunal report and extract sections.

        Supports cross-talk: can import T131 (psychiatric), T134 (nursing),
        or T133 (social circumstances) reports and map by heading content.
        """
        from docx import Document
        import re

        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"[SOCIAL] Failed to open DOCX: {e}")
            return {}

        # Detect report type
        full_text = ' '.join([p.text for p in doc.paragraphs]).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += ' ' + cell.text.lower()

        if 'social' in full_text and 'circumstances' in full_text:
            form_type = 'T133'
        elif 'nursing' in full_text:
            form_type = 'T134'
        elif 'responsible clinician' in full_text or 'medical' in full_text:
            form_type = 'T131'
        else:
            form_type = 'unknown'

        print(f"[SOCIAL] Detected report type: {form_type}")
        result = {'form_type': form_type, 'sections': {}}

        # Helper to get unique cells (merged cells repeat same content)
        def get_unique_cells(cell_list):
            unique = []
            for c in cell_list:
                if c and (not unique or c != unique[-1]):
                    unique.append(c)
            return unique

        def is_question_text(text):
            """Check if text is a question heading, not answer content."""
            if not text or len(text) < 10:
                return False
            lower = text.lower().strip()
            lower_clean = re.sub(r'^[\[\]☐☒xX\s\-–\d\.]*', '', lower).strip()
            question_patterns = [
                'are there any factors', 'are there any adjustments',
                'what is the nature of', 'give details of any',
                'what are the strengths', 'give a summary of',
                'in section 2 cases', 'in all other cases',
                'if the patient was discharged', 'if the patient were discharged',
                'please explain how', 'is there any other relevant',
                'do you have any recommendations', 'is the patient now suffering',
                'what appropriate and available', 'what are the dates',
                'what are the circumstances', 'give reasons for',
                'does the patient have a learning', 'what is it about',
                'would they be likely to act', 'managed effectively in the community',
            ]
            for pattern in question_patterns:
                if lower_clean.startswith(pattern):
                    return True
            return False

        # Parse tables for section content using HEADING-BASED matching
        for table in doc.tables:
            rows = list(table.rows)
            if not rows:
                continue

            # --- Only match heading on FIRST ROW of each table ---
            first_cells = [cell.text.strip() for cell in rows[0].cells]
            if not any(first_cells):
                # Fall through to special-case handling below
                pass
            else:
                unique_first = get_unique_cells(first_cells)

                heading_key = None
                heading_cell_idx = -1
                for cell_idx, cell_text in enumerate(unique_first):
                    if re.match(r'^\d{1,2}\.\s*$', cell_text):
                        continue
                    matched = self._match_heading_to_card(cell_text)
                    if matched:
                        heading_key = matched
                        heading_cell_idx = cell_idx
                        break

                if heading_key:
                    answer = ""

                    # Check cells after heading in first row
                    start_idx = heading_cell_idx + 1 if heading_cell_idx >= 0 else 1
                    if len(unique_first) > start_idx:
                        for cell_text in unique_first[start_idx:]:
                            if is_question_text(cell_text):
                                continue
                            if cell_text.strip() in ('No\nYes', 'Yes\nNo', 'No', 'Yes', 'N/A'):
                                continue
                            if cell_text and not re.match(r'^\d+\.\s*', cell_text):
                                cleaned = self._clean_docx_content(cell_text)
                                if cleaned and cleaned not in ('No', 'Yes', 'N/A'):
                                    answer = cleaned
                                    break

                    # Scan ALL remaining rows for answer (multi-row Yes/No tables)
                    if not answer:
                        yes_checked = False
                        no_checked = False

                        for look_ahead in range(1, len(rows)):
                            row_cells = [cell.text.strip() for cell in rows[look_ahead].cells]
                            unique_row = get_unique_cells(row_cells)

                            for idx, cell_text in enumerate(unique_row):
                                stripped = cell_text.strip()
                                if not stripped:
                                    continue

                                # Detect Yes/No checkbox state from cells
                                if stripped in ('No', 'Yes', 'N/A'):
                                    continue
                                if stripped in ('No\nYes', 'Yes\nNo'):
                                    continue

                                # Check for "x" or "X" marker next to Yes/No
                                if stripped.lower() in ('x', 'x ', ' x'):
                                    # Check what the previous cell was (Yes or No)
                                    if idx > 0:
                                        prev = unique_row[idx - 1].strip()
                                        if prev == 'Yes':
                                            yes_checked = True
                                        elif prev == 'No':
                                            no_checked = True
                                    continue

                                # "Yes  X" or "Yes X" inline
                                if re.match(r'^Yes\s+[xX]', stripped):
                                    yes_checked = True
                                    continue
                                if re.match(r'^No\s+[xX]', stripped):
                                    no_checked = True
                                    continue

                                # "X- no formal employment" style (checkbox with note)
                                x_note = re.match(r'^[xX]\s*[-–]\s*(.+)', stripped)
                                if x_note:
                                    # This is a checkbox note, check if it's for Yes or No
                                    if idx > 0 and unique_row[idx - 1].strip() == 'No':
                                        no_checked = True
                                    elif idx > 0 and unique_row[idx - 1].strip() == 'Yes':
                                        yes_checked = True
                                    continue

                                # Skip sub-questions ("If Yes, what are they?")
                                if is_question_text(cell_text):
                                    continue
                                lower_stripped = stripped.lower()
                                if lower_stripped.startswith('if yes') or lower_stripped.startswith('if no'):
                                    continue
                                if lower_stripped.startswith('what is the name of') or lower_stripped.startswith('what are they'):
                                    continue
                                if lower_stripped.startswith('by what date'):
                                    continue

                                # Check for checkbox format
                                if '[x' in cell_text.lower() or '☒' in cell_text:
                                    yes_answer = self._extract_checkbox_answer(cell_text)
                                    if yes_answer:
                                        answer = yes_answer
                                        break

                                inline_answer = self._parse_inline_checkbox(cell_text)
                                if inline_answer:
                                    if inline_answer == 'Yes':
                                        yes_checked = True
                                    elif inline_answer == 'No':
                                        no_checked = True
                                    else:
                                        answer = inline_answer
                                    break

                                # Substantial text = actual answer content
                                cleaned = self._clean_docx_content(cell_text)
                                if cleaned and len(cleaned) > 5 and not is_question_text(cleaned):
                                    if cleaned.replace('\n', ' ').strip() not in ('No Yes', 'Yes No', 'No', 'Yes', 'N/A'):
                                        answer = cleaned
                                        break

                            if answer:
                                break

                        # For Yes/No-only sections with no explanation text
                        if not answer:
                            if yes_checked:
                                answer = "Yes"
                            elif no_checked:
                                answer = "No"

                    if answer:
                        # Add Yes prefix for Yes/No sections
                        yes_no_keys = {'factors_hearing', 'adjustments', 's2_detention',
                                      'other_detention', 'discharge_risk', 'recommendations',
                                      'mappa', 'care_plan_funding', 'employment'}
                        if heading_key in yes_no_keys and answer not in ('Yes', 'No', 'N/A'):
                            if not answer.startswith('Yes') and not answer.startswith('No'):
                                answer = f"Yes - {answer}"
                        result['sections'][heading_key] = answer
                        print(f"[SOCIAL] Matched heading ({heading_key}): {answer[:50]}...")

                    continue  # Done with this table

            # --- Special-case handling for tables without heading matches ---
            # Iterate rows for patient details, author info, signature
            for i in range(len(rows)):
                cells = [cell.text.strip() for cell in rows[i].cells]
                if not any(cells):
                    continue

                unique_cells = get_unique_cells(cells)
                first_cell = cells[0] if cells else ''
                first_lower = first_cell.lower()

                # Patient details table - various field name formats
                if 'name of patient' in first_lower or 'full name' in first_lower:
                    if len(unique_cells) > 1:
                        result['sections']['patient_details'] = f"Name: {unique_cells[1]}"
                    elif i + 1 < len(rows):
                        next_cells = [c.text.strip() for c in rows[i + 1].cells]
                        next_unique = get_unique_cells(next_cells)
                        if next_unique and next_unique[0]:
                            result['sections']['patient_details'] = f"Name: {next_unique[0]}"
                elif 'date of birth' in first_lower or first_lower.startswith('dob'):
                    dob_val = unique_cells[1] if len(unique_cells) > 1 else ""
                    if not dob_val and i + 1 < len(rows):
                        next_cells = [c.text.strip() for c in rows[i + 1].cells]
                        next_unique = get_unique_cells(next_cells)
                        dob_val = next_unique[0] if next_unique else ""
                    if dob_val:
                        if 'patient_details' not in result['sections']:
                            result['sections']['patient_details'] = ""
                        result['sections']['patient_details'] += f"\nDOB: {dob_val}"
                elif 'nhs number' in first_lower:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nNHS: {unique_cells[1]}"
                elif 'usual place of residence' in first_lower or 'usual place' in first_lower:
                    addr_val = unique_cells[1] if len(unique_cells) > 1 else ""
                    if not addr_val and i + 1 < len(rows):
                        next_cells = [c.text.strip() for c in rows[i + 1].cells]
                        next_unique = get_unique_cells(next_cells)
                        addr_val = next_unique[0] if next_unique else ""
                    if addr_val and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nAddress: {addr_val}"
                elif 'mental health act status' in first_lower:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nMHA Status: {unique_cells[1]}"

                # Author info
                elif 'your name' in first_lower:
                    author_val = unique_cells[1] if len(unique_cells) > 1 else ""
                    if not author_val and i + 1 < len(rows):
                        next_cells = [c.text.strip() for c in rows[i + 1].cells]
                        next_unique = get_unique_cells(next_cells)
                        author_val = next_unique[0] if next_unique else ""
                    if author_val:
                        result['sections']['author'] = f"Name: {author_val}"
                elif 'your role' in first_lower:
                    role_val = unique_cells[1] if len(unique_cells) > 1 else ""
                    if not role_val and i + 1 < len(rows):
                        next_cells = [c.text.strip() for c in rows[i + 1].cells]
                        next_unique = get_unique_cells(next_cells)
                        role_val = next_unique[0] if next_unique else ""
                    if role_val and 'author' in result['sections']:
                        result['sections']['author'] += f"\nRole: {role_val}"
                elif 'date of report' in first_lower:
                    date_val = unique_cells[1] if len(unique_cells) > 1 else ""
                    if not date_val and i + 1 < len(rows):
                        next_cells = [c.text.strip() for c in rows[i + 1].cells]
                        next_unique = get_unique_cells(next_cells)
                        date_val = next_unique[0] if next_unique else ""
                    if date_val and 'author' in result['sections']:
                        result['sections']['author'] += f"\nDate: {date_val}"

                # Signature table
                elif first_cell in ('Signed:', 'Signature'):
                    if len(unique_cells) > 1:
                        result['sections']['signature'] = f"Signed: {unique_cells[1]}"
                elif first_cell == 'Print Name:' and 'signature' in result['sections']:
                    if len(unique_cells) > 1:
                        result['sections']['signature'] += f"\nName: {unique_cells[1]}"
                elif first_cell == 'Date:' and 'signature' in result['sections']:
                    if len(unique_cells) > 1:
                        result['sections']['signature'] += f"\nDate: {unique_cells[1]}"

        print(f"[SOCIAL] Parsed DOCX: found {len(result['sections'])} sections")
        for key in result['sections']:
            print(f"  - {key}")
        return result

    def _clean_docx_content(self, text: str) -> str:
        """Clean up content from DOCX, removing checkbox formatting."""
        import re

        if not text:
            return ""

        # Remove checkbox patterns like [   ], [ X ], [x], etc.
        text = re.sub(r'\[\s*[xX]?\s*\]', '', text)

        # Remove "No" and "Yes" on their own lines
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line.lower() in ('no', 'yes', 'n/a'):
                continue
            # Remove various "If yes" patterns at start of line
            line = re.sub(r'^[–-]\s*If yes[^:]*:\s*', '', line)
            line = re.sub(r'^If yes[^:]*:\s*', '', line)
            if line:
                cleaned_lines.append(line)

        result = '\n'.join(cleaned_lines).strip()
        return result

    def _parse_inline_checkbox(self, text: str) -> str:
        """Parse inline checkbox format like 'No [ ] Yes [X] N/A [ ]' to extract checked value."""
        import re

        if not text:
            return ""

        # Look for checked boxes [x], [X], [ x], [x ], [ X ], etc.
        checked_pattern = r'(No|Yes|N/A)\s*\[\s*[xX]\s*\]'
        match = re.search(checked_pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)

        return ""

    def _extract_checkbox_answer(self, text: str) -> str:
        """Extract answer from checkbox format like [x] - If yes, explain..."""
        import re
        # Look for checked checkbox followed by content
        match = re.search(r'\[[\s]*[xX][\s]*\][^a-zA-Z]*(?:If yes[^\n]*\n+)?(.+)', text, re.DOTALL | re.IGNORECASE)
        if match:
            content = match.group(1).strip()
            if content and len(content) > 5:
                return content
        # Also check for ☒
        match = re.search(r'☒[^a-zA-Z]*(?:If yes[^\n]*\n+)?(.+)', text, re.DOTALL | re.IGNORECASE)
        if match:
            content = match.group(1).strip()
            if content and len(content) > 5:
                return content
        return ""

    def _extract_yes_no_answer(self, cell_text: str, next_cell: str = "") -> tuple:
        """Extract answer from Yes/No checkbox format. Returns (answer, has_explanation)."""
        import re

        # Check for inline format
        inline_result = self._parse_inline_checkbox(cell_text)
        if inline_result:
            return inline_result, False

        # Check for two-cell format where first cell is "No\nYes" and second has checkboxes
        if cell_text.startswith(('No\n', 'Yes\n')) and next_cell:
            lines = next_cell.split('\n')

            # Find which checkbox is checked (has X)
            no_checked = False
            yes_checked = False
            for i, line in enumerate(lines):
                if re.search(r'\[\s*[xX]\s*\]', line):
                    if i == 0:
                        no_checked = True
                    else:
                        yes_checked = True
                        break

            if no_checked and not yes_checked:
                return "No", False

            if yes_checked:
                # Look for explanation after "If yes"
                if 'If yes' in next_cell:
                    match = re.search(r'If yes[^:?]*[:\?]\s*\n+(.*)', next_cell, re.DOTALL)
                    if match:
                        explanation = match.group(1).strip()
                        if explanation:
                            return explanation, True
                    parts = next_cell.split('\n\n')
                    if len(parts) > 1:
                        explanation = parts[1].strip()
                        if explanation:
                            return explanation, True
                else:
                    # No "If yes" - look for explanation directly after checkbox
                    explanation_lines = []
                    found_checkbox = False
                    for line in lines:
                        if found_checkbox and line.strip():
                            explanation_lines.append(line.strip())
                        if re.search(r'\[\s*[xX]\s*\]', line):
                            found_checkbox = True
                    if explanation_lines:
                        explanation = '\n'.join(explanation_lines)
                        return explanation, True
                return "Yes", False

        # Check for checked box in the cell itself
        if '[x' in cell_text.lower():
            lines = cell_text.split('\n')
            for i, line in enumerate(lines):
                if '[x' in line.lower() or '[ x' in line.lower():
                    if i > 0 and lines[i-1].strip().lower() in ('no', 'yes'):
                        return lines[i-1].strip(), False

        return "", False

    def _populate_from_docx(self, result: dict, file_path: str):
        """Populate report sections from parsed DOCX tribunal form data."""
        from PySide6.QtWidgets import QMessageBox
        from shared_data_store import get_shared_store
        import os

        sections = result.get('sections', {})

        # Store imported report data for _has_report_data() and popup population
        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        filename = os.path.basename(file_path)
        action = self._ask_import_action(filename, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        merged_sections = {}
        for key, content in sections.items():
            if content:
                if action == "add":
                    content = self._merge_report_section(key, content, filename)
                self._imported_report_data[key] = content
                merged_sections[key] = content

        # Cards stay empty until user clicks card heading (trigger)
        loaded_count = len(merged_sections)

        # Push sections to shared store for cross-talk with other forms
        shared_store = get_shared_store()
        shared_store.set_report_sections(merged_sections, source_form="social_tribunal")

        # Show success message
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"Successfully {action_word.lower()} social circumstances report from:\n{filename}\n\n"
            f"{action_word} {loaded_count} sections.\n\n"
            f"Review and edit the content as needed."
        )

        print(f"[SOCIAL] {action_word} {loaded_count} sections from DOCX")

        # Populate popups with imported data (Yes/No states, details, collapsible sections)
        self._populate_popups_with_imported_data(merged_sections)

    def _send_to_data_extractor(self, file_path: str):
        """Send a file to the data extractor for processing."""
        self._data_extractor_source_file = file_path
        self._open_data_extractor_overlay()

        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            if hasattr(self._data_extractor_overlay, 'load_file'):
                self._data_extractor_overlay.load_file(file_path)
            elif hasattr(self._data_extractor_overlay, 'upload_and_extract'):
                self._data_extractor_overlay.upload_and_extract()
            print(f"[SOCIAL] Sent to data extractor: {file_path}")

    def _ask_import_action(self, source_filename: str, import_type: str = "report") -> str:
        """Ask user whether to add to existing imported data or replace it.

        Returns: 'add', 'replace', or 'cancel'
        """
        from PySide6.QtWidgets import QMessageBox

        if import_type == "report":
            has_existing = hasattr(self, '_imported_report_data') and any(self._imported_report_data.values())
        else:
            has_existing = hasattr(self, '_extracted_raw_notes') and bool(self._extracted_raw_notes)

        if not has_existing:
            return "replace"

        msg = QMessageBox(self)
        msg.setWindowTitle("Import Data")
        if import_type == "report":
            msg.setText(
                f"Report data has already been imported.\n\n"
                f"New file: {source_filename}\n\n"
                f"Would you like to add this data to the existing import, or replace it?"
            )
        else:
            msg.setText(
                f"Clinical notes have already been loaded.\n\n"
                f"Would you like to add these notes to the existing set, or replace them?"
            )
        add_btn = msg.addButton("Add to Existing", QMessageBox.AcceptRole)
        replace_btn = msg.addButton("Replace All", QMessageBox.DestructiveRole)
        cancel_btn = msg.addButton("Cancel", QMessageBox.RejectRole)
        msg.setDefaultButton(add_btn)
        msg.exec()

        clicked = msg.clickedButton()
        if clicked == cancel_btn:
            return "cancel"
        elif clicked == replace_btn:
            return "replace"
        return "add"

    def _merge_report_section(self, key: str, new_content: str, source_filename: str) -> str:
        """Merge new report content with existing section, adding source reference."""
        existing = self._imported_report_data.get(key, "")
        if not existing:
            return new_content
        return f"{existing}\n\n─── Source: {source_filename} ───\n\n{new_content}"

    def _populate_from_pdf(self, result: dict, file_path: str):
        """Populate report sections from parsed PDF tribunal form data."""
        from PySide6.QtWidgets import QMessageBox
        from pdf_loader import format_radio_value
        import os

        form_type = result.get('form_type', 'unknown')
        sections = result.get('sections', {})

        # Map section keys to card keys for social report
        section_to_card = {
            'patient_details': 'patient_details',
            'author': 'author',
            'factors_hearing': 'factors_hearing',
            'adjustments': 'adjustments',
            'forensic': 'forensic',
            'previous_mh_dates': 'previous_mh',
            'diagnosis': 'diagnosis',
            'treatment': 'treatment',
            'strengths': 'strengths',
            'progress': 'progress',
            'compliance': 'compliance',
            'risk_harm': 'risk_harm',
            'risk_property': 'risk_property',
            'discharge_risk': 'discharge_risk',
            'community': 'community',
            'recommendations': 'recommendations',
            'signature_date': 'signature',
        }

        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        filename = os.path.basename(file_path)
        action = self._ask_import_action(filename, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        # Map and merge sections
        merged_sections = {}
        for section_key, content in sections.items():
            card_key = section_to_card.get(section_key)
            if card_key and content:
                if action == "add":
                    content = self._merge_report_section(card_key, content, filename)
                self._imported_report_data[card_key] = content
                merged_sections[card_key] = content

        loaded_count = len(merged_sections)

        # Push patient details to shared store for other forms
        self._push_patient_details_to_shared_store(sections)

        # Show success message
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "PDF Loaded",
            f"Successfully {action_word.lower()} {form_type} report from:\n{filename}\n\n"
            f"{action_word} {loaded_count} sections from the PDF.\n\n"
            f"Review and edit the content as needed."
        )

        print(f"[PDF] {action_word} {loaded_count} sections from {form_type} report")

        # Populate popups with imported data collapsible sections
        self._populate_popups_with_imported_data(merged_sections)

    def _populate_popups_with_imported_data(self, sections: dict):
        """Populate popups with imported data - creates popups and calls _populate_single_popup."""
        for section_key, content in sections.items():
            if not content or not content.strip():
                continue

            # Create popup if it doesn't exist (popups are normally created lazily)
            if section_key not in self.popups:
                popup = self._create_popup(section_key)
                if popup:
                    self.popups[section_key] = popup
                    self.popup_stack.addWidget(popup)
                    if hasattr(popup, 'sent'):
                        popup.sent.connect(lambda text, k=section_key: self._update_card(k, text))
                    print(f"[SOCIAL] Created popup '{section_key}' during import")

            if section_key in self.popups:
                popup = self.popups[section_key]
                self._populate_single_popup(popup, section_key, content)

    def _populate_single_popup(self, popup, section_key: str, content: str):
        """Populate a single popup with imported content - copied from psychiatric tribunal."""
        import re

        # ============================================================
        # HELPER FUNCTIONS FOR CHECKBOX DETECTION
        # ============================================================
        def has_yes_cross(text):
            """Check if Yes has a checked box symbol."""
            if re.search(r'yes\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☒☑✓✔\[xX\]].*no\s*[☐□\[\s\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes.*[☐□]\s*no', text, re.IGNORECASE):
                return True
            return False

        def has_no_cross(text):
            """Check if No has a checked box symbol (meaning answer is No)."""
            if re.search(r'no\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            if re.search(r'no\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*no\b', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☐□\[\s\]].*no\s*[☒☑✓✔\[xX\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☐□]\s*yes.*[☒☑✓✔]\s*no', text, re.IGNORECASE):
                return True
            return False

        def is_yes_content(text):
            """Check if content indicates Yes answer."""
            if has_yes_cross(text):
                return True
            if has_no_cross(text):
                return False
            lower = text.lower().strip()
            if lower.startswith("yes") or "yes -" in lower or "yes," in lower:
                return True
            # If substantial text and not explicitly "no", assume yes
            if len(text.strip()) > 30 and not lower.startswith("no"):
                return True
            return False

        def extract_detail(text):
            """Extract detail text after Yes/No prefix."""
            for prefix in ["Yes - ", "Yes, ", "Yes-", "yes - ", "yes, ", "yes-", "Yes ", "yes "]:
                if text.startswith(prefix):
                    return text[len(prefix):].strip()
            return text

        # ============================================================
        # PATIENT DETAILS - parse text into fields directly
        # ============================================================
        if section_key == "patient_details":
            from datetime import datetime
            patient_info = {}

            # Parse "Key: Value" lines from imported text
            name_match = re.search(r"(?:Full\s*Name|Name|Patient)[:\s]+(.+)", content, re.IGNORECASE)
            if name_match:
                patient_info["name"] = name_match.group(1).strip()

            dob_match = re.search(r"(?:Date of Birth|DOB|D\.O\.B)[:\s]+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})", content, re.IGNORECASE)
            if dob_match:
                dob_str = dob_match.group(1).strip()
                for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y"]:
                    try:
                        patient_info["dob"] = datetime.strptime(dob_str, fmt)
                        break
                    except ValueError:
                        continue

            # Try text date: "7 October 1979" or "18th July, 1991"
            if not patient_info.get("dob"):
                text_dob = re.search(
                    r"(?:Date of Birth|DOB|D\.O\.B)[:\s]+(\d{1,2})(?:st|nd|rd|th)?\s*,?\s*(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s*,?\s*(\d{4})",
                    content, re.IGNORECASE
                )
                if text_dob:
                    month_map = {'january':1,'jan':1,'february':2,'feb':2,'march':3,'mar':3,'april':4,'apr':4,'may':5,'june':6,'jun':6,'july':7,'jul':7,'august':8,'aug':8,'september':9,'sep':9,'october':10,'oct':10,'november':11,'nov':11,'december':12,'dec':12}
                    try:
                        patient_info["dob"] = datetime(int(text_dob.group(3)), month_map.get(text_dob.group(2).lower(), 1), int(text_dob.group(1)))
                    except ValueError:
                        pass

            gender_match = re.search(r"(?:Gender|Sex)[:\s]*(Male|Female|M|F)\b", content, re.IGNORECASE)
            if gender_match:
                g = gender_match.group(1).upper()
                patient_info["gender"] = "Male" if g in ("MALE", "M") else "Female"

            address_match = re.search(r"(?:Usual Place of Residence|Address|Residence)[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
            if address_match:
                patient_info["address"] = address_match.group(1).strip()

            if patient_info and hasattr(popup, 'fill_patient_info'):
                popup.fill_patient_info(patient_info)
                print(f"[SOCIAL] Filled patient_details fields: {list(patient_info.keys())}")
            return

        # ============================================================
        # YES/NO POPUPS WITH STANDARD yes_btn/no_btn
        # ============================================================
        yes_no_sections = ["factors_hearing", "adjustments", "compliance",
                          "discharge_risk", "recommendations"]

        if section_key in yes_no_sections:
            is_yes = is_yes_content(content)
            print(f"[SOCIAL] Section '{section_key}' content starts: {repr(content[:100])}")
            print(f"[SOCIAL] Section '{section_key}' is_yes={is_yes}")

            if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
                if is_yes:
                    popup.yes_btn.setChecked(True)
                    print(f"[SOCIAL] Set '{section_key}' popup to Yes")
                else:
                    popup.no_btn.setChecked(True)
                    print(f"[SOCIAL] Set '{section_key}' popup to No")

                # For factors_hearing, also detect and set the specific factor radio buttons
                if section_key == "factors_hearing" and is_yes:
                    content_lower = content.lower()
                    if hasattr(popup, 'autism_rb') and ('autism' in content_lower or 'autistic' in content_lower):
                        popup.autism_rb.setChecked(True)
                        print(f"[SOCIAL] Set factors_hearing to Autism")
                    elif hasattr(popup, 'ld_rb') and ('learning disability' in content_lower or 'learning difficulties' in content_lower):
                        popup.ld_rb.setChecked(True)
                        print(f"[SOCIAL] Set factors_hearing to Learning Disability")
                    elif hasattr(popup, 'patience_rb') and ('irritab' in content_lower or 'frustration' in content_lower or 'patience' in content_lower):
                        popup.patience_rb.setChecked(True)
                        print(f"[SOCIAL] Set factors_hearing to Low frustration tolerance")

        # ============================================================
        # EMPLOYMENT POPUP - click Yes and fill details_edit with narrative
        # ============================================================
        if section_key in ("employment", "care_plan_funding") and hasattr(popup, 'yes_btn') and hasattr(popup, 'details_edit'):
            is_yes = is_yes_content(content)
            if is_yes:
                popup.yes_btn.setChecked(True)
                popup.no_btn.setChecked(False)
                if hasattr(popup, 'details_container'):
                    popup.details_container.show()
                print(f"[SOCIAL] Set employment popup to Yes")
            else:
                popup.no_btn.setChecked(True)
                popup.yes_btn.setChecked(False)
                print(f"[SOCIAL] Set employment popup to No")

        # ============================================================
        # MAPPA POPUP - detect yes/no and auto-click
        # ============================================================
        if section_key == "mappa" and hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
            is_yes = is_yes_content(content)
            print(f"[SOCIAL] MAPPA content starts: {repr(content[:100])}")
            print(f"[SOCIAL] MAPPA is_yes={is_yes}")
            if is_yes:
                popup.yes_btn.setChecked(True)
                popup.no_btn.setChecked(False)
                if hasattr(popup, 'details_container'):
                    popup.details_container.show()
                # Try to detect MAPPA level from content
                content_lower = content.lower()
                if hasattr(popup, 'level_combo'):
                    if 'level 3' in content_lower or 'enhanced' in content_lower:
                        popup.level_combo.setCurrentIndex(3)
                    elif 'level 2' in content_lower or 'active multi' in content_lower:
                        popup.level_combo.setCurrentIndex(2)
                    elif 'level 1' in content_lower or 'ordinary' in content_lower:
                        popup.level_combo.setCurrentIndex(1)
                # Try to detect MAPPA category from content
                if hasattr(popup, 'category_combo'):
                    if 'category 4' in content_lower or 'terrorism' in content_lower:
                        popup.category_combo.setCurrentIndex(4)
                    elif 'category 3' in content_lower or 'other dangerous' in content_lower:
                        popup.category_combo.setCurrentIndex(3)
                    elif 'category 2' in content_lower or 'violent' in content_lower:
                        popup.category_combo.setCurrentIndex(2)
                    elif 'category 1' in content_lower or 'sex offender' in content_lower:
                        popup.category_combo.setCurrentIndex(1)
                print(f"[SOCIAL] Set MAPPA popup to Yes")
            else:
                popup.no_btn.setChecked(True)
                popup.yes_btn.setChecked(False)
                print(f"[SOCIAL] Set MAPPA popup to No")
            self._add_imported_data_to_popup(popup, section_key, content)
            return

        # For popups with built-in imported data section (set_entries), use that
        # Skip set_entries for yes_no sections (already handled above with radio buttons/details)
        # Also skip set_entries when importing report data - use amber CollapsibleSection instead
        has_report_data = hasattr(self, '_imported_report_data') and bool(self._imported_report_data)
        if not has_report_data and section_key not in yes_no_sections and hasattr(popup, 'set_entries') and callable(popup.set_entries):
            if isinstance(content, str):
                # Convert string content to entries list format (set_entries expects list of dicts)
                paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
                if not paragraphs and content.strip():
                    paragraphs = [content.strip()]
                entries_list = [{"date": None, "text": p} for p in paragraphs]
                popup.set_entries(entries_list)
            else:
                popup.set_entries(content)
            print(f"[SOCIAL] Called set_entries for '{section_key}'")
            # Check if set_entries actually populated visible content;
            # if not (e.g. stub set_entries), fall through to collapsible section
            has_visual = (
                (hasattr(popup, '_extracted_checkboxes') and popup._extracted_checkboxes) or
                (hasattr(popup, 'extracted_section') and hasattr(popup.extracted_section, 'isVisible') and popup.extracted_section.isVisible())
            )
            if has_visual:
                popup._imported_data_added = True
            else:
                self._add_imported_data_to_popup(popup, section_key, content)
        else:
            # Add collapsible imported data section for other popup types
            self._add_imported_data_to_popup(popup, section_key, content)

    def _add_imported_data_to_popup(self, popup, section_key: str, content: str):
        """Add imported data collapsible section to popup with checkboxes."""
        from PySide6.QtWidgets import QLabel, QVBoxLayout, QWidget, QCheckBox, QHBoxLayout, QFrame, QSizePolicy
        from PySide6.QtCore import Qt
        import re

        # Skip if empty or minimal checkbox-only content
        if not content or not content.strip():
            return
        cleaned = content.strip()
        if re.match(r'^(No|Yes)\s*\[\s*\]\s*(No|Yes)\s*\[\s*\]$', cleaned, re.IGNORECASE):
            return
        if re.match(r'^[☐☒\s]*(No|Yes)\s*[☐☒\s]*(No|Yes)\s*[☐☒\s]*$', cleaned, re.IGNORECASE):
            return
        if cleaned.lower() in ('yes', 'no', 'n/a', 'yes.', 'no.'):
            return

        # Skip sections that don't need imported data collapsible
        skip_sections = {'s2_detention', 'other_detention', 'author', 'signature', 'patient_details', 'factors_hearing', 'adjustments'}
        if section_key in skip_sections:
            return

        # Skip if popup already has imported data
        if getattr(popup, '_imported_data_added', False):
            return

        # Try to find target layout - more robust approach
        target_layout = None
        parent_widget = None
        is_simple_popup = False  # True if popup has no scroll area (simple QWidget)

        if hasattr(popup, 'scroll_layout') and popup.scroll_layout:
            target_layout = popup.scroll_layout
            parent_widget = target_layout.parentWidget()
        elif hasattr(popup, 'main_layout') and popup.main_layout:
            target_layout = popup.main_layout
            parent_widget = target_layout.parentWidget()
        elif hasattr(popup, 'container_layout') and popup.container_layout:
            target_layout = popup.container_layout
            parent_widget = target_layout.parentWidget()
        else:
            # QWidget.layout() returns the top-level layout
            # Some popups shadow self.layout with a QVBoxLayout instance, so protect against TypeError
            try:
                layout_obj = popup.layout()
            except TypeError:
                layout_obj = None
            if layout_obj and hasattr(layout_obj, 'insertWidget'):
                target_layout = layout_obj
                parent_widget = popup
                is_simple_popup = True  # No scroll area - start collapsed to avoid overlap

        if not target_layout or not callable(getattr(target_layout, 'insertWidget', None)):
            print(f"[SOCIAL] No target layout found for '{section_key}'")
            return

        if not parent_widget:
            parent_widget = popup

        # For simple popups, wrap all existing content in a QScrollArea so the
        # CollapsibleSection and input widgets coexist without overlap
        if is_simple_popup:
            from PySide6.QtWidgets import QScrollArea as _ScrollArea
            scroll = _ScrollArea()
            scroll.setWidgetResizable(True)
            scroll.setFrameShape(_ScrollArea.Shape.NoFrame)
            scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            inner = QWidget()
            inner.setStyleSheet("background: transparent;")
            inner_layout = QVBoxLayout(inner)
            inner_layout.setContentsMargins(target_layout.contentsMargins())
            inner_layout.setSpacing(target_layout.spacing())

            # Move existing widgets from popup layout into inner container
            while target_layout.count():
                item = target_layout.takeAt(0)
                w = item.widget()
                if w:
                    inner_layout.addWidget(w)
                elif item.layout():
                    inner_layout.addLayout(item.layout())
                else:
                    inner_layout.addItem(item)

            scroll.setWidget(inner)
            target_layout.setContentsMargins(0, 0, 0, 0)
            target_layout.addWidget(scroll)

            # Use inner layout as the target for CollapsibleSection insertion
            target_layout = inner_layout
            parent_widget = inner
            print(f"[SOCIAL] Wrapped simple popup '{section_key}' in QScrollArea")

        try:
            from background_history_popup import CollapsibleSection

            import_section = CollapsibleSection("Imported Data", parent=parent_widget, start_collapsed=False)
            import_section.set_content_height(150)
            import_section._min_height = 80
            import_section._max_height = 2000
            import_section.set_header_style("""
                QFrame {
                    background: rgba(254, 243, 199, 0.95);
                    border: 1px solid rgba(245, 158, 11, 0.5);
                    border-radius: 6px 6px 0 0;
                }
            """)
            import_section.title_label.setStyleSheet("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #d97706;
                    background: transparent;
                    border: none;
                }
            """)

            # Create content widget
            content_widget = QWidget()
            content_widget.setStyleSheet("""
                QWidget {
                    background: rgba(254, 243, 199, 0.95);
                    border: 1px solid rgba(245, 158, 11, 0.4);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            content_layout = QVBoxLayout(content_widget)
            content_layout.setContentsMargins(12, 10, 12, 10)
            content_layout.setSpacing(4)

            # Get current card content to check what's already there
            card_text = ""
            if section_key in self.cards:
                card = self.cards[section_key]
                if hasattr(card, 'editor'):
                    card_text = card.editor.toPlainText().lower()

            # Single checkbox + scrollable QTextEdit for ALL imported data
            from PySide6.QtWidgets import QTextEdit as _QTextEdit
            cleaned_content = content.strip()

            # Scale height based on content length
            est_lines = max(8, len(cleaned_content) // 55 + 2)
            text_height = min(600, est_lines * 18)
            section_height = text_height + 40

            import_section.set_content_height(section_height)
            import_section._max_height = section_height + 100

            cb = QCheckBox()
            cb.setFixedSize(20, 20)
            cb.setStyleSheet("QCheckBox::indicator { width: 16px; height: 16px; }")
            cb.setProperty("full_text", cleaned_content)
            cb.setProperty("section_key", section_key)

            cb.setChecked(False)
            cb.toggled.connect(lambda checked, sk=section_key, txt=cleaned_content: self._on_imported_checkbox_toggled(sk, txt, checked))

            row = QHBoxLayout()
            row.setContentsMargins(0, 4, 0, 4)
            row.setSpacing(8)
            row.setAlignment(Qt.AlignmentFlag.AlignTop)
            row.addWidget(cb, 0, Qt.AlignmentFlag.AlignTop)

            text_edit = _QTextEdit()
            text_edit.setPlainText(cleaned_content)
            text_edit.setReadOnly(True)
            text_edit.setMinimumHeight(text_height)
            text_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
            text_edit.setStyleSheet("""
                QTextEdit {
                    font-size: 17px; color: #4a4a4a;
                    background: transparent; border: none;
                }
            """)
            row.addWidget(text_edit, 1)

            row_widget = QFrame()
            row_widget.setStyleSheet("QFrame { background: transparent; border: none; }")
            row_widget.setLayout(row)
            content_layout.addWidget(row_widget)

            popup._imported_checkboxes = [cb]

            import_section.set_content(content_widget)

            target_layout.insertWidget(0, import_section)
            target_layout.addStretch()
            popup._imported_data_added = True
            print(f"[SOCIAL] Added imported data for '{section_key}'")

        except Exception as e:
            print(f"[SOCIAL] Error adding imported data to '{section_key}': {e}")

    def _on_imported_checkbox_toggled(self, section_key: str, text: str, checked: bool):
        """Handle when an imported data checkbox is toggled."""
        if section_key not in self.cards:
            return

        card = self.cards[section_key]
        if not hasattr(card, 'editor'):
            return

        # Block textChanged signals to prevent _on_card_text_changed from
        # syncing back to popup and re-triggering _send_to_card
        card.editor.blockSignals(True)
        try:
            current_text = card.editor.toPlainText()

            if checked:
                # Add text to card if not already there
                if text not in current_text:
                    if current_text.strip():
                        new_text = current_text.strip() + "\n\n" + text
                    else:
                        new_text = text
                    card.editor.setPlainText(new_text)
                    print(f"[SOCIAL] Added imported text to '{section_key}'")
            else:
                # Remove imported text from card
                if text in current_text:
                    new_text = current_text.replace(text, "").strip()
                    while "\n\n\n" in new_text:
                        new_text = new_text.replace("\n\n\n", "\n\n")
                    card.editor.setPlainText(new_text)
                    print(f"[SOCIAL] Removed imported text from '{section_key}'")
                else:
                    # Fallback: clear card and regenerate from popup form only
                    card.editor.setPlainText("")
                    print(f"[SOCIAL] Cleared card for '{section_key}'")
        finally:
            card.editor.blockSignals(False)

    def _push_patient_details_to_shared_store(self, sections: dict):
        """Extract patient details from loaded sections and push to shared store."""
        import re
        from shared_data_store import get_shared_store

        patient_details = sections.get('patient_details', '')
        if not patient_details:
            return

        patient_info = {
            "name": None,
            "dob": None,
            "nhs_number": None,
            "gender": None,
            "age": None,
            "ethnicity": None,
        }

        from datetime import datetime

        # Extract name
        name_match = re.search(r"(?:Name|Patient)[:\s]+([A-Za-z][A-Za-z\-\' ]+)", patient_details, re.IGNORECASE)
        if name_match:
            patient_info["name"] = name_match.group(1).strip()

        # Extract DOB
        dob_match = re.search(r"(?:DOB|Date of Birth|D\.O\.B)[:\s]+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})", patient_details, re.IGNORECASE)
        if dob_match:
            dob_str = dob_match.group(1).strip()
            for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y"]:
                try:
                    patient_info["dob"] = datetime.strptime(dob_str, fmt)
                    break
                except ValueError:
                    continue

        # Extract NHS number
        nhs_match = re.search(r"(?:NHS)[:\s]*(\d{3}\s*\d{3}\s*\d{4}|\d{10})", patient_details, re.IGNORECASE)
        if nhs_match:
            nhs = nhs_match.group(1).replace(" ", "")
            patient_info["nhs_number"] = f"{nhs[:3]} {nhs[3:6]} {nhs[6:]}" if len(nhs) == 10 else nhs

        # Extract gender
        gender_match = re.search(r"(?:Gender|Sex)[:\s]*(Male|Female|M|F)\b", patient_details, re.IGNORECASE)
        if gender_match:
            g = gender_match.group(1).upper()
            patient_info["gender"] = "Male" if g in ("MALE", "M") else "Female" if g in ("FEMALE", "F") else None

        # Extract age
        age_patterns = [
            r"(?:AGE)[:\s]*(\d{1,3})\s*(?:years?|yrs?|y\.?o\.?)?\b",
            r"\b(\d{1,3})\s*(?:year|yr)\s*old\b",
            r"\b(\d{1,3})\s*y\.?o\.?\b",
            r"\baged?\s*(\d{1,3})\b",
        ]
        for pattern in age_patterns:
            match = re.search(pattern, patient_details, re.IGNORECASE)
            if match:
                age_val = int(match.group(1))
                if 0 < age_val < 120:
                    patient_info["age"] = age_val
                    break

        # Calculate age from DOB if not found explicitly
        if not patient_info["age"] and patient_info["dob"]:
            today = datetime.today()
            dob = patient_info["dob"]
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if 0 < age < 120:
                patient_info["age"] = age

        # Extract ethnicity
        ethnicity_patterns = [
            r"(?:ETHNICITY|ETHNIC\s*(?:GROUP|ORIGIN)?)[:\s]+([A-Za-z][A-Za-z\s\-\/]+?)(?:\n|$|,)",
            r"\b(White\s*(?:British|Irish|European|Other)?)\b",
            r"\b(Black\s*(?:British|African|Caribbean|Other)?)\b",
            r"\b(Asian\s*(?:British|Indian|Pakistani|Bangladeshi|Chinese|Other)?)\b",
            r"\b(Mixed\s*(?:White\s*(?:and|&)\s*(?:Black\s*(?:Caribbean|African)|Asian))?)\b",
        ]
        for pattern in ethnicity_patterns:
            match = re.search(pattern, patient_details, re.IGNORECASE)
            if match:
                ethnicity_val = match.group(1).strip()
                if len(ethnicity_val) > 2:
                    patient_info["ethnicity"] = ethnicity_val.title()
                    break

        # Push to shared store if any info found
        if any(patient_info.values()):
            shared_store = get_shared_store()
            shared_store.set_patient_info(patient_info, source="social_tribunal")
            print(f"[SOCIAL] Pushed patient details to SharedDataStore: {list(k for k,v in patient_info.items() if v)}")

    def _open_data_extractor_overlay(self):
        """Create the data extractor (hidden) for background processing."""
        from data_extractor_popup import DataExtractorPopup

        if not hasattr(self, '_data_extractor_overlay') or not self._data_extractor_overlay:
            self._data_extractor_overlay = DataExtractorPopup(parent=self)
            self._data_extractor_overlay.hide()
            if hasattr(self._data_extractor_overlay, 'data_extracted'):
                self._data_extractor_overlay.data_extracted.connect(self._on_data_extracted)

    # Mapping from PTR card keys to SCT card keys for report imports
    # PTR1->SCT1, PTR2->SCT2, PTR3->SCT3, PTR5->SCT4, PTR6->SCT5,
    # PTR13->SCT15, PTR14->SCT16, PTR17->SCT17, PTR18->SCT18,
    # PTR21->SCT27, PTR22->SCT28
    PTR_TO_SCT_MAP = {
        "patient_details": "patient_details",      # PTR1  -> SCT1
        "factors_hearing": "factors_hearing",      # PTR3  -> SCT2
        "adjustments": "adjustments",              # PTR4  -> SCT3
        "forensic": "forensic",                    # PTR5  -> SCT4
        "previous_mh_dates": "previous_mh_dates",  # PTR6  -> SCT5
        "strengths": "strengths",                  # PTR13 -> SCT15
        "progress": "progress",                    # PTR14 -> SCT16
        "risk_harm": "risk_harm",                  # PTR17 -> SCT17
        "risk_property": "risk_property",          # PTR18 -> SCT18
        "discharge_risk": "discharge_risk",        # PTR21 -> SCT27
        "community": "community",                  # PTR22 -> SCT28
    }

    def _on_data_extracted(self, data: dict):
        import os
        print(f"[SOCIAL] Data extracted: {list(data.get('categories', {}).keys())}")
        cov = data.get("_coverage")
        if cov and cov.get("uncategorised", 0) > 0:
            print(f"[SOCIAL] Warning: {cov['uncategorised']} paragraphs uncategorised "
                  f"({cov['categorised']}/{cov['total_paragraphs']} categorised)")

        if 'filtered_category' in data and self._selected_card_key:
            filtered_category = data['filtered_category']
            categories = data.get('categories', {})
            for cat_name, cat_data in categories.items():
                if filtered_category.lower() in cat_name.lower():
                    entries = cat_data.get('entries', [])
                    content_parts = []
                    for entry in entries[:10]:
                        content = entry.get('content', '') or entry.get('text', '')
                        if content:
                            content_parts.append(content[:300])
                    if content_parts:
                        self._update_card(self._selected_card_key, '\n\n'.join(content_parts))
            return

        # Skip if this exact data was already processed
        categories = data.get('categories', {})
        cat_keys = tuple(sorted(categories.keys())) if categories else ()
        cat_count = sum(len(v.get("items", [])) if isinstance(v, dict) else 0 for v in categories.values())
        content_sig = (cat_keys, cat_count)
        if self._data_processed_id == content_sig:
            print(f"[SOCIAL] Skipping _on_data_extracted - data already processed")
            return
        self._data_processed_id = content_sig

        # Check if data came from a report (not notes)
        is_report = False
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            dtype = getattr(self._data_extractor_overlay, '_detected_document_type', None)
            if dtype == "reports":
                is_report = True
                print(f"[SOCIAL] Detected report data (dtype={dtype})")

        categories = data.get('categories', {})
        print(f"[SOCIAL] Available categories: {list(categories.keys())}")

        if is_report and categories:
            # Report pipeline: map PTR categories to SCT sections
            source = os.path.basename(getattr(self, '_data_extractor_source_file', '') or '') or "Data Extractor"
            self._populate_from_report_categories(categories, source_filename=source)
        else:
            # Notes pipeline - skip if report data already imported (prevents cross-talk)
            if self._has_report_data():
                print(f"[SOCIAL] Skipping notes pipeline - report data already imported")
                return

            raw_notes = []
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                raw_notes = getattr(self._data_extractor_overlay, 'notes', [])

            if not raw_notes:
                try:
                    from shared_data_store import get_shared_store
                    shared_store = get_shared_store()
                    if shared_store.has_notes():
                        raw_notes = shared_store.notes
                        print(f"[SOCIAL] Got {len(raw_notes)} notes from SharedDataStore")
                except Exception as e:
                    print(f"[SOCIAL] Error getting notes from SharedDataStore: {e}")

            # Ask add/replace if existing notes
            action = self._ask_import_action("", "notes")
            if action == "cancel":
                return
            if action == "add":
                existing_notes = getattr(self, '_extracted_raw_notes', []) or []
                raw_notes = existing_notes + raw_notes
                existing_cats = getattr(self, '_extracted_categories', {}) or {}
                for cat_name, cat_data in categories.items():
                    if cat_name in existing_cats and isinstance(existing_cats[cat_name], dict) and isinstance(cat_data, dict):
                        existing_items = existing_cats[cat_name].get("items", [])
                        new_items = cat_data.get("items", [])
                        existing_cats[cat_name]["items"] = existing_items + new_items
                    else:
                        existing_cats[cat_name] = cat_data
                categories = existing_cats

            self._extracted_categories = categories
            self._extracted_raw_notes = raw_notes
            self._incident_data = raw_notes

            print(f"[SOCIAL] Raw notes with datetime dates: {len(raw_notes)}")

            self._populate_fixed_panels(categories)

    def _populate_from_report_categories(self, categories: dict, source_filename: str = ""):
        """Populate SCT sections from PTR report categories.

        Maps PTR card keys to SCT card keys using PTR_TO_SCT_MAP and
        populates card editors + popups with imported data.
        """
        from PySide6.QtWidgets import QMessageBox
        from shared_data_store import get_shared_store

        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}
        if not hasattr(self, '_imported_report_sections'):
            self._imported_report_sections = {}

        source_label = source_filename or "Data Extractor"
        action = self._ask_import_action(source_label, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()
            self._imported_report_sections.clear()

        valid_sct_keys = {key for _, key in self.SECTIONS}

        card_texts = {}
        all_ptr_texts = {}  # Original PTR keys for SharedDataStore cross-talk

        for cat_name, cat_data in categories.items():
            # Combine all items' text first (needed for both SCT and cross-talk)
            items = cat_data.get("items", []) if isinstance(cat_data, dict) else []
            texts = []
            for item in items:
                if isinstance(item, dict):
                    text = item.get("text", "") or item.get("content", "")
                elif isinstance(item, str):
                    text = item
                else:
                    text = str(item)
                if text:
                    texts.append(text.strip())

            if not texts:
                continue

            combined = "\n\n".join(texts)

            # Always store under original PTR key for cross-talk
            if cat_name in all_ptr_texts:
                all_ptr_texts[cat_name] += "\n\n" + combined
            else:
                all_ptr_texts[cat_name] = combined

            # Map PTR card key to SCT card key for local population
            if cat_name in self.PTR_TO_SCT_MAP:
                sct_key = self.PTR_TO_SCT_MAP[cat_name]
            elif cat_name in valid_sct_keys:
                sct_key = cat_name
            else:
                print(f"[SOCIAL] No SCT mapping for PTR category: {cat_name} (will pass through to other forms)")
                continue

            if sct_key in card_texts:
                card_texts[sct_key] += "\n\n" + combined
            else:
                card_texts[sct_key] = combined

            if sct_key not in self._imported_report_sections:
                self._imported_report_sections[sct_key] = []
            self._imported_report_sections[sct_key].append((cat_name, combined))

        sections_for_store = {}
        for sct_key, content in card_texts.items():
            if action == "add":
                content = self._merge_report_section(sct_key, content, source_label)
            self._imported_report_data[sct_key] = content
            sections_for_store[sct_key] = content
            print(f"[SOCIAL] Stored report data for SCT section '{sct_key}'")

        # Populate popups with imported data
        self._populate_popups_with_imported_data(sections_for_store)

        # Push patient details to shared store
        self._push_patient_details_to_shared_store(sections_for_store)

        # Push ALL original PTR keys to shared store for cross-talk
        # This ensures PTR/NTR get sections that don't map to SCT
        shared_store = get_shared_store()
        cross_talk_sections = dict(all_ptr_texts)
        cross_talk_sections.update(sections_for_store)
        shared_store.set_report_sections(cross_talk_sections, source_form="social_tribunal")

        mapped_count = len(sections_for_store)
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"{action_word} report data from {source_label}.\n\n"
            f"{action_word} {mapped_count} PTR sections to SCT sections.\n\n"
            f"Click each card to review and send the content."
        )
        print(f"[SOCIAL] {action_word} {mapped_count} PTR sections to SCT sections")

    def set_notes(self, notes: list):
        """
        Set notes from shared data store.
        Called by MainWindow when notes are available from another section.
        """
        if not notes:
            return

        # Skip if report data exists (report takes priority over notes)
        if self._has_report_data():
            print(f"[Social] Skipping set_notes - report data already imported")
            return

        # Skip if these exact notes were already processed
        notes_sig = (len(notes), id(notes))
        if self._notes_processed_id == notes_sig:
            print(f"[Social] Skipping set_notes - notes already processed")
            return
        self._notes_processed_id = notes_sig

        # Store raw notes at page level for use in sections
        self._extracted_raw_notes = notes

        # If data extractor exists, update its notes too
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            if hasattr(self._data_extractor_overlay, 'set_notes'):
                self._data_extractor_overlay.set_notes(notes)

        print(f"[Social] Received {len(notes)} notes from shared store")

    def _populate_fixed_panels(self, categories: dict):
        """Populate all fixed panels with extracted data (matches medical report approach)."""
        raw_notes = self._extracted_raw_notes

        if not raw_notes:
            print("[SOCIAL] No extracted data to populate panels")
            return

        # Helper to get items from category (categories have "items" key)
        def get_category_items(cat_name):
            cat = categories.get(cat_name, {})
            if isinstance(cat, dict):
                return cat.get("items", [])
            return []

        # Section 5: Past psychiatric history
        psych_items = get_category_items("Past Psychiatric History")
        psych_entries = psych_items if psych_items else raw_notes[:20]
        print(f"[SOCIAL] Section 5: Found {len(psych_items)} Past Psychiatric History items, using {len(psych_entries)} entries")

        if "previous_mh_dates" in self.popups:
            popup = self.popups["previous_mh_dates"]
            if hasattr(popup, 'notes'):
                popup.notes = raw_notes
            if hasattr(popup, 'set_entries'):
                popup.set_entries(psych_entries, f"{len(psych_entries)} notes")
                print(f"[SOCIAL] Populated section 5 with {len(psych_entries)} notes")
            # Run timeline analysis to detect admissions
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[SOCIAL] Ran timeline analysis for section 5")

        # Store for pending (when popup created on click) - always store with fallback
        if not hasattr(self, '_pending_section_data'):
            self._pending_section_data = {}
        self._pending_section_data['previous_mh_dates'] = psych_entries
        self._pending_section_data['previous_community'] = psych_entries
        print(f"[SOCIAL] Stored pending data for previous_mh_dates: {len(psych_entries)} entries")

        # Sections 6, 7, 8: Background History (home_family, housing, financial)
        background_items = get_category_items("Background History")
        if not background_items:
            background_items = get_category_items("background history")
        if not background_items:
            background_items = get_category_items("Personal History")
        if not background_items:
            # Fallback: use raw notes if no background category
            background_items = raw_notes[:20] if raw_notes else []
            print(f"[SOCIAL] Sections 6-8: No Background History category, using {len(background_items)} raw notes as fallback")
        background_entries = background_items if background_items else []
        print(f"[SOCIAL] Sections 6-8: Found {len(background_entries)} Background History entries")

        # Store for sections 6, 7, 8
        if background_entries:
            self._pending_section_data['home_family'] = background_entries
            self._pending_section_data['housing'] = background_entries
            self._pending_section_data['financial'] = background_entries
            self._pending_section_data['employment'] = background_entries
            self._pending_section_data['nearest_relative'] = background_entries
            print(f"[SOCIAL] Stored pending data for home_family, housing, financial: {len(background_entries)} entries")

        # Section 16: Progress - Use notes from last 6 months (identical to medical section 14)
        from datetime import datetime, timedelta

        def parse_note_date_s16(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # Find the most recent note date
        all_note_dates_s16 = []
        for n in raw_notes:
            note_date = parse_note_date_s16(n.get('date') or n.get('datetime'))
            if note_date:
                all_note_dates_s16.append(note_date)

        if all_note_dates_s16:
            most_recent_date_s16 = max(all_note_dates_s16)
            oldest_date_s16 = min(all_note_dates_s16)
            # 6 months before the most recent entry
            six_months_cutoff_s16 = most_recent_date_s16 - timedelta(days=180)

            print(f"[SOCIAL] Section 16 DEBUG: Most recent note: {most_recent_date_s16.strftime('%d/%m/%Y')}")
            print(f"[SOCIAL] Section 16 DEBUG: Oldest note: {oldest_date_s16.strftime('%d/%m/%Y')}")
            print(f"[SOCIAL] Section 16 DEBUG: 6-month cutoff: {six_months_cutoff_s16.strftime('%d/%m/%Y')}")

            # Filter notes from last 6 months (relative to most recent entry)
            notes_with_dates_s16 = []
            for n in raw_notes:
                note_date = parse_note_date_s16(n.get('date') or n.get('datetime'))
                if note_date and note_date >= six_months_cutoff_s16:
                    notes_with_dates_s16.append(n)

            # Sort by date (most recent first)
            recent_progress = sorted(
                notes_with_dates_s16,
                key=lambda x: parse_note_date_s16(x.get('date') or x.get('datetime')),
                reverse=True
            )

            if recent_progress:
                filtered_dates = [parse_note_date_s16(n.get('date') or n.get('datetime')) for n in recent_progress]
                filtered_dates = [d for d in filtered_dates if d]
                if filtered_dates:
                    print(f"[SOCIAL] Section 16 DEBUG: Filtered range: {min(filtered_dates).strftime('%d/%m/%Y')} to {max(filtered_dates).strftime('%d/%m/%Y')}")
        else:
            # Fallback: no parseable dates, use first 100 raw notes
            recent_progress = raw_notes[:100]

        print(f"[SOCIAL] Section 16: Prepared {len(recent_progress)} notes from last 6 months (total raw: {len(raw_notes)})")

        if "progress" in self.popups:
            popup = self.popups["progress"]
            if hasattr(popup, 'notes'):
                popup.notes = raw_notes
            if hasattr(popup, 'set_entries'):
                popup.set_entries(recent_progress, f"{len(recent_progress)} notes (last 6 months)")
                print(f"[SOCIAL] Populated section 16 with {len(recent_progress)} notes from last 6 months")

        # Store for pending
        if not hasattr(self, '_pending_section_data'):
            self._pending_section_data = {}
        self._pending_section_data['progress'] = recent_progress

        # Section 17 & 18: Incidents - search all notes using incidentDICT (identical to medical)
        from pathlib import Path
        import re

        # Load incident terms from dictionary file
        incident_terms = []
        incident_file = Path(__file__).parent / "incidentDICT.txt"
        if incident_file.exists():
            with open(incident_file, 'r', encoding='utf-8') as f:
                for line in f:
                    term = line.strip().lower()
                    if term:
                        incident_terms.append(term)
            print(f"[SOCIAL] Loaded {len(incident_terms)} incident terms from incidentDICT.txt")

        exclude_keywords = ['h/o', 'history of', 'previous noted', 'previously noted', 'previous history',
                            'risk of', 'historical risk', 'past risk',
                            'no aggression', 'nil aggression', 'nil aggressive', 'not aggressive',
                            'no aggressive', 'low risk and nil', 'no risk', 'nil on the shift',
                            'fire safety', 'fire test', 'gas / fire', 'gas/fire',
                            'risk and aggression:', 'risk/aggression', 'risk & aggression',
                            'politely declined', 'hearing 30th', 'managers hearing', 'manangers hearing',
                            'solicitor today', 'no signs of aggression', 'no physical or verbal aggression',
                            'did not display any aggressive', 'no aggressive behaviour',
                            'no challenging or aggressive', 'no episode of aggressive',
                            'no irritable or aggressive', 'no physical/verbal aggression',
                            'no instances of irritability, aggressive',
                            'did not appear intoxicated', 'not intoxicated', 'no intoxication',
                            'although did not appear intox', 'did not appear to be intox',
                            'if no further abusive', 'no further abusive']

        def is_sexual_health_or_history(text):
            pattern1 = r'\bsexual\s+(?:\w+\s+)?(?:health|history|hx)\b'
            pattern2 = r'\bsexually\s+transmitted\b'
            return bool(re.search(pattern1, text, re.IGNORECASE) or re.search(pattern2, text, re.IGNORECASE))

        def parse_note_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # Extract individual incident lines
        incident_lines = []
        for note in raw_notes:
            content = note.get('content', '') or note.get('text', '') or ''
            content = content.replace('\r\n', '\n').replace('\r', '\n')

            date_val = note.get('date') or note.get('datetime')
            date_obj = parse_note_date(date_val)
            date_str = date_obj.strftime('%d/%m/%Y') if date_obj else 'Unknown'

            all_segments = []
            for line in content.split('\n'):
                sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', line)
                all_segments.extend(sentences)

            for idx, line in enumerate(all_segments):
                line_clean = line.strip()
                line_clean = ''.join(c for c in line_clean if ord(c) >= 32 or c == '\t')
                line_lower = line_clean.lower()
                if not line_clean or len(line_clean) < 10:
                    continue

                if any(ex in line_lower for ex in exclude_keywords):
                    continue

                if is_sexual_health_or_history(line_lower):
                    continue

                word_boundary_terms = {'anger'}

                for term in incident_terms:
                    if term in word_boundary_terms:
                        pattern = r'\b' + re.escape(term) + r'\b'
                        if not re.search(pattern, line_lower):
                            continue
                    else:
                        if term not in line_lower:
                            continue

                    final_text = line_clean
                    if 'datix' in line_lower or 'ulysses' in line_lower:
                        extra_lines = []
                        for extra_idx in range(1, 4):
                            if idx + extra_idx < len(all_segments):
                                extra_line = all_segments[idx + extra_idx].strip()
                                extra_line = ''.join(c for c in extra_line if ord(c) >= 32 or c == '\t')
                                if extra_line and len(extra_line) > 5:
                                    extra_lines.append(extra_line)
                        if extra_lines:
                            final_text = line_clean + ' | ' + ' | '.join(extra_lines)

                    incident_lines.append({
                        'date': date_str,
                        'date_obj': date_obj,
                        'text': f"{date_str}: {final_text}",
                        'content': f"{date_str}: {final_text}",
                        'term': term
                    })
                    break

        # Sort by date (most recent first)
        incident_lines.sort(key=lambda x: x['date_obj'] or datetime.min, reverse=True)

        # Deduplicate by line content
        seen_lines = set()
        incidents = []
        for inc in incident_lines:
            line_key = inc['text'].lower()
            if line_key not in seen_lines:
                seen_lines.add(line_key)
                incidents.append(inc)

        print(f"[SOCIAL] Found {len(incidents)} unique incident lines")

        incidents_filtered = incidents
        self._incident_data = incidents_filtered

        # Use same raw notes for both sections 17 and 18 (TribunalRiskHarmPopup/TribunalRiskPropertyPopup)
        if "risk_harm" in self.popups:
            popup = self.popups["risk_harm"]
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[SOCIAL] Populated section 17 (risk_harm) with {len(raw_notes)} notes for risk analysis")
            elif hasattr(popup, 'set_entries'):
                popup.set_entries(incidents_filtered, f"{len(incidents_filtered)} incidents")
                print(f"[SOCIAL] Populated section 17 with {len(incidents_filtered)} incidents")

        if "risk_property" in self.popups:
            popup = self.popups["risk_property"]
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[SOCIAL] Populated section 18 (risk_property) with {len(raw_notes)} notes for risk analysis")
            elif hasattr(popup, 'set_entries'):
                popup.set_entries(incidents_filtered, f"{len(incidents_filtered)} incidents")
                print(f"[SOCIAL] Populated section 18 with {len(incidents_filtered)} incidents")

        # Store raw notes for pending (used by TribunalRiskHarmPopup/TribunalRiskPropertyPopup)
        self._pending_section_data['risk_harm'] = raw_notes
        self._pending_section_data['risk_property'] = raw_notes

        # Section 4: Forensic History - populate forensic history panel
        forensic_items = get_category_items("Forensic History")
        if not forensic_items:
            forensic_items = get_category_items("forensic history")
        if not forensic_items:
            forensic_items = get_category_items("FORENSIC")

        if "forensic" in self.popups:
            popup = self.popups["forensic"]
            # Get notes for risk analysis (incidents - Physical Aggression, Property Damage, Sexual Behaviour)
            notes = getattr(self, '_extracted_raw_notes', [])
            if hasattr(popup, 'set_forensic_data'):
                popup.set_forensic_data(notes, forensic_items if forensic_items else [])
                print(f"[SOCIAL] Populated section 4 forensic panel with {len(notes)} notes and {len(forensic_items) if forensic_items else 0} entries")

        # Store forensic items for later use (when popup is created)
        if not hasattr(self, '_pending_forensic_data'):
            self._pending_forensic_data = {}
        self._pending_forensic_data['forensic'] = forensic_items if forensic_items else []
        if forensic_items:
            self._pending_section_data['forensic'] = forensic_items
        print(f"[SOCIAL] Stored {len(forensic_items) if forensic_items else 0} forensic entries for forensic popup")

        print("[SOCIAL] Fixed panels populated with extracted data")


# ================================================================
# SOCIAL-SPECIFIC POPUPS
# ================================================================

class HomeFamilyPopup(QWidget):
    """Popup for home and family circumstances with collapsible imported data - sends to card on click."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._entries = []
        self._extracted_checkboxes = []
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        from background_history_popup import CollapsibleSection

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # IMPORTED DATA SECTION
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 60
        self.extracted_section._max_height = 2000
        self.extracted_section.set_header_style("""
            QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 6px 6px 0 0; }
        """)
        self.extracted_section.set_title_style("""
            QLabel { font-size: 18px; font-weight: 600; color: #806000; background: transparent; border: none; }
        """)

        extracted_content = QWidget()
        extracted_content.setStyleSheet("""
            QWidget { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-top: none; border-radius: 0 0 12px 12px; }
        """)

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _send_to_card(self):
        """Send checked items directly to card (replaces card content)."""
        parts = []
        for cb in self._extracted_checkboxes:
            if cb.isChecked():
                parts.append(cb.property("full_text"))
        self.sent.emit("\n\n".join(parts))

    def set_entries(self, entries: list, date_range_info: str = ""):
        self._entries = entries

        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if isinstance(entries, str):
            entries = [{"date": None, "text": p.strip()} for p in entries.split("\n\n") if p.strip()] if entries.strip() else []

        if entries:
            def get_sort_date(item):
                dt = item.get("date")
                return dt.strftime("%Y-%m-%d") if dt and hasattr(dt, "strftime") else str(dt) if dt else ""

            sorted_items = sorted(entries, key=get_sort_date, reverse=True)

            for item in sorted_items:
                dt = item.get("date")
                text = (item.get("text", "") or item.get("content", "")).strip()
                if not text:
                    continue

                date_str = dt.strftime("%d %b %Y") if dt and hasattr(dt, "strftime") else str(dt) if dt else "No date"

                entry_frame = QFrame()
                entry_frame.setObjectName("entryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                entry_frame.setStyleSheet("QFrame#entryFrame { background: rgba(255,255,255,0.95); border: 1px solid rgba(180,150,50,0.4); border-radius: 8px; padding: 4px; }")
                entry_layout = QVBoxLayout(entry_frame)
                entry_layout.setContentsMargins(10, 8, 10, 8)
                entry_layout.setSpacing(6)
                entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                # Toggle button on LEFT with arrow
                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet("QPushButton { background: rgba(180,150,50,0.2); border: none; border-radius: 4px; font-size: 17px; font-weight: bold; color: #806000; } QPushButton:hover { background: rgba(180,150,50,0.35); }")
                header_row.addWidget(toggle_btn)

                date_label = QLabel(f"{date_str}")
                date_label.setStyleSheet("QLabel { font-size: 17px; font-weight: 600; color: #806000; background: transparent; border: none; }")
                date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                header_row.addWidget(date_label)
                header_row.addStretch()

                # Checkbox on RIGHT
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.stateChanged.connect(self._send_to_card)
                header_row.addWidget(cb)

                entry_layout.addLayout(header_row)

                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setStyleSheet("QTextEdit { font-size: 17px; color: #333; background: rgba(255,248,220,0.5); border: none; padding: 8px; border-radius: 6px; }")
                body_text.setMinimumHeight(60)
                body_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                body_text.setVisible(False)
                entry_layout.addWidget(body_text)

                drag_bar = QFrame()
                drag_bar.setFixedHeight(8)
                drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
                drag_bar.setStyleSheet("""
                    QFrame {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                        border-radius: 2px; margin: 2px 40px;
                    }
                    QFrame:hover {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                    }
                """)
                drag_bar.setVisible(False)
                drag_bar._drag_y = None
                drag_bar._init_h = None
                def _make_drag_handlers(handle, text_widget):
                    def press(ev):
                        handle._drag_y = ev.globalPosition().y()
                        handle._init_h = text_widget.height()
                    def move(ev):
                        if handle._drag_y is not None:
                            delta = int(ev.globalPosition().y() - handle._drag_y)
                            new_h = max(60, handle._init_h + delta)
                            text_widget.setMinimumHeight(new_h)
                            text_widget.setMaximumHeight(new_h)
                    def release(ev):
                        if handle._drag_y is not None:
                            text_widget.setMaximumHeight(16777215)
                            handle._drag_y = None
                    return press, move, release
                dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
                drag_bar.mousePressEvent = dp
                drag_bar.mouseMoveEvent = dm
                drag_bar.mouseReleaseEvent = dr
                entry_layout.addWidget(drag_bar)

                def make_toggle(btn, body, frame, popup_self, bar):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        bar.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                        popup_self.extracted_container.updateGeometry()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
                toggle_btn.clicked.connect(toggle_fn)
                date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                self.extracted_checkboxes_layout.addWidget(entry_frame)
                self._extracted_checkboxes.append(cb)

            self.extracted_checkboxes_layout.addStretch()
            self.extracted_section.setVisible(True)
            # Keep collapsed on open - user can expand manually
        else:
            self.extracted_section.setVisible(False)


class HousingPopup(QWidget):
    """Popup for housing/accommodation with collapsible imported data - sends to card on click."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._entries = []
        self._extracted_checkboxes = []
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        from background_history_popup import CollapsibleSection

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # IMPORTED DATA SECTION
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 60
        self.extracted_section._max_height = 2000
        self.extracted_section.set_header_style("QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 6px 6px 0 0; }")
        self.extracted_section.set_title_style("QLabel { font-size: 18px; font-weight: 600; color: #806000; background: transparent; border: none; }")

        extracted_content = QWidget()
        extracted_content.setStyleSheet("QWidget { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-top: none; border-radius: 0 0 12px 12px; }")

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _send_to_card(self):
        """Send checked items directly to card (replaces card content)."""
        parts = [cb.property("full_text") for cb in self._extracted_checkboxes if cb.isChecked()]
        self.sent.emit("\n\n".join(parts))

    def set_entries(self, entries: list, date_range_info: str = ""):
        self._entries = entries

        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if isinstance(entries, str):
            entries = [{"date": None, "text": p.strip()} for p in entries.split("\n\n") if p.strip()] if entries.strip() else []

        if entries:
            def get_sort_date(item):
                dt = item.get("date")
                return dt.strftime("%Y-%m-%d") if dt and hasattr(dt, "strftime") else str(dt) if dt else ""

            sorted_items = sorted(entries, key=get_sort_date, reverse=True)

            for item in sorted_items:
                dt = item.get("date")
                text = (item.get("text", "") or item.get("content", "")).strip()
                if not text:
                    continue

                date_str = dt.strftime("%d %b %Y") if dt and hasattr(dt, "strftime") else str(dt) if dt else "No date"

                entry_frame = QFrame()
                entry_frame.setObjectName("entryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                entry_frame.setStyleSheet("QFrame#entryFrame { background: rgba(255,255,255,0.95); border: 1px solid rgba(180,150,50,0.4); border-radius: 8px; padding: 4px; }")
                entry_layout = QVBoxLayout(entry_frame)
                entry_layout.setContentsMargins(10, 8, 10, 8)
                entry_layout.setSpacing(6)
                entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                # Toggle button on LEFT with arrow
                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet("QPushButton { background: rgba(180,150,50,0.2); border: none; border-radius: 4px; font-size: 17px; font-weight: bold; color: #806000; } QPushButton:hover { background: rgba(180,150,50,0.35); }")
                header_row.addWidget(toggle_btn)

                date_label = QLabel(f"{date_str}")
                date_label.setStyleSheet("QLabel { font-size: 17px; font-weight: 600; color: #806000; background: transparent; border: none; }")
                date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                header_row.addWidget(date_label)
                header_row.addStretch()

                # Checkbox on RIGHT
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.stateChanged.connect(self._send_to_card)
                header_row.addWidget(cb)

                entry_layout.addLayout(header_row)

                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setStyleSheet("QTextEdit { font-size: 17px; color: #333; background: rgba(255,248,220,0.5); border: none; padding: 8px; border-radius: 6px; }")
                body_text.setMinimumHeight(60)
                body_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                body_text.setVisible(False)
                entry_layout.addWidget(body_text)

                drag_bar = QFrame()
                drag_bar.setFixedHeight(8)
                drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
                drag_bar.setStyleSheet("""
                    QFrame {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                        border-radius: 2px; margin: 2px 40px;
                    }
                    QFrame:hover {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                    }
                """)
                drag_bar.setVisible(False)
                drag_bar._drag_y = None
                drag_bar._init_h = None
                def _make_drag_handlers(handle, text_widget):
                    def press(ev):
                        handle._drag_y = ev.globalPosition().y()
                        handle._init_h = text_widget.height()
                    def move(ev):
                        if handle._drag_y is not None:
                            delta = int(ev.globalPosition().y() - handle._drag_y)
                            new_h = max(60, handle._init_h + delta)
                            text_widget.setMinimumHeight(new_h)
                            text_widget.setMaximumHeight(new_h)
                    def release(ev):
                        if handle._drag_y is not None:
                            text_widget.setMaximumHeight(16777215)
                            handle._drag_y = None
                    return press, move, release
                dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
                drag_bar.mousePressEvent = dp
                drag_bar.mouseMoveEvent = dm
                drag_bar.mouseReleaseEvent = dr
                entry_layout.addWidget(drag_bar)

                def make_toggle(btn, body, frame, popup_self, bar):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        bar.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                        popup_self.extracted_container.updateGeometry()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
                toggle_btn.clicked.connect(toggle_fn)
                date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                self.extracted_checkboxes_layout.addWidget(entry_frame)
                self._extracted_checkboxes.append(cb)

            self.extracted_checkboxes_layout.addStretch()
            self.extracted_section.setVisible(True)
            # Keep collapsed on open - user can expand manually
        else:
            self.extracted_section.setVisible(False)


class FinancialPopup(QWidget):
    """Popup for financial position with collapsible imported data - sends to card on click."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._entries = []
        self._extracted_checkboxes = []
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        from background_history_popup import CollapsibleSection

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # IMPORTED DATA SECTION
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 60
        self.extracted_section._max_height = 2000
        self.extracted_section.set_header_style("QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 6px 6px 0 0; }")
        self.extracted_section.set_title_style("QLabel { font-size: 18px; font-weight: 600; color: #806000; background: transparent; border: none; }")

        extracted_content = QWidget()
        extracted_content.setStyleSheet("QWidget { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-top: none; border-radius: 0 0 12px 12px; }")

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _send_to_card(self):
        """Send checked items directly to card (replaces card content)."""
        parts = [cb.property("full_text") for cb in self._extracted_checkboxes if cb.isChecked()]
        self.sent.emit("\n\n".join(parts))

    def set_entries(self, entries: list, date_range_info: str = ""):
        self._entries = entries

        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if isinstance(entries, str):
            entries = [{"date": None, "text": p.strip()} for p in entries.split("\n\n") if p.strip()] if entries.strip() else []

        if entries:
            def get_sort_date(item):
                dt = item.get("date")
                return dt.strftime("%Y-%m-%d") if dt and hasattr(dt, "strftime") else str(dt) if dt else ""

            sorted_items = sorted(entries, key=get_sort_date, reverse=True)

            for item in sorted_items:
                dt = item.get("date")
                text = (item.get("text", "") or item.get("content", "")).strip()
                if not text:
                    continue

                date_str = dt.strftime("%d %b %Y") if dt and hasattr(dt, "strftime") else str(dt) if dt else "No date"

                entry_frame = QFrame()
                entry_frame.setObjectName("entryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                entry_frame.setStyleSheet("QFrame#entryFrame { background: rgba(255,255,255,0.95); border: 1px solid rgba(180,150,50,0.4); border-radius: 8px; padding: 4px; }")
                entry_layout = QVBoxLayout(entry_frame)
                entry_layout.setContentsMargins(10, 8, 10, 8)
                entry_layout.setSpacing(6)
                entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                # Toggle button on LEFT with arrow
                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet("QPushButton { background: rgba(180,150,50,0.2); border: none; border-radius: 4px; font-size: 17px; font-weight: bold; color: #806000; } QPushButton:hover { background: rgba(180,150,50,0.35); }")
                header_row.addWidget(toggle_btn)

                date_label = QLabel(f"{date_str}")
                date_label.setStyleSheet("QLabel { font-size: 17px; font-weight: 600; color: #806000; background: transparent; border: none; }")
                date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                header_row.addWidget(date_label)
                header_row.addStretch()

                # Checkbox on RIGHT
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.stateChanged.connect(self._send_to_card)
                header_row.addWidget(cb)

                entry_layout.addLayout(header_row)

                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setStyleSheet("QTextEdit { font-size: 17px; color: #333; background: rgba(255,248,220,0.5); border: none; padding: 8px; border-radius: 6px; }")
                body_text.setMinimumHeight(60)
                body_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                body_text.setVisible(False)
                entry_layout.addWidget(body_text)

                drag_bar = QFrame()
                drag_bar.setFixedHeight(8)
                drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
                drag_bar.setStyleSheet("""
                    QFrame {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                        border-radius: 2px; margin: 2px 40px;
                    }
                    QFrame:hover {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                    }
                """)
                drag_bar.setVisible(False)
                drag_bar._drag_y = None
                drag_bar._init_h = None
                def _make_drag_handlers(handle, text_widget):
                    def press(ev):
                        handle._drag_y = ev.globalPosition().y()
                        handle._init_h = text_widget.height()
                    def move(ev):
                        if handle._drag_y is not None:
                            delta = int(ev.globalPosition().y() - handle._drag_y)
                            new_h = max(60, handle._init_h + delta)
                            text_widget.setMinimumHeight(new_h)
                            text_widget.setMaximumHeight(new_h)
                    def release(ev):
                        if handle._drag_y is not None:
                            text_widget.setMaximumHeight(16777215)
                            handle._drag_y = None
                    return press, move, release
                dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
                drag_bar.mousePressEvent = dp
                drag_bar.mouseMoveEvent = dm
                drag_bar.mouseReleaseEvent = dr
                entry_layout.addWidget(drag_bar)

                def make_toggle(btn, body, frame, popup_self, bar):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        bar.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                        popup_self.extracted_container.updateGeometry()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
                toggle_btn.clicked.connect(toggle_fn)
                date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                self.extracted_checkboxes_layout.addWidget(entry_frame)
                self._extracted_checkboxes.append(cb)

            self.extracted_checkboxes_layout.addStretch()
            self.extracted_section.setVisible(True)
            # Keep collapsed on open - user can expand manually
        else:
            self.extracted_section.setVisible(False)


class EmploymentPopup(QWidget):
    """Popup for employment opportunities with collapsible imported data - sends to card on click."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._entries = []
        self._extracted_checkboxes = []
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        from background_history_popup import CollapsibleSection

        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        main_scroll = QScrollArea()
        main_scroll.setWidgetResizable(True)
        main_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        main_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        main_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)

        main_container = QWidget()
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(4, 4, 4, 4)
        main_layout.setSpacing(8)

        # ========== INPUT SECTION ==========
        input_frame = QFrame()
        input_frame.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 8px; }")
        input_layout = QVBoxLayout(input_frame)
        input_layout.setContentsMargins(16, 16, 16, 16)
        input_layout.setSpacing(12)

        # Question label
        label = QLabel("Are there opportunities for employment if discharged?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        input_layout.addWidget(label)

        # Yes/No buttons
        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes")
        self.no_btn = QPushButton("No")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6; border: 1px solid #d1d5db;
                    padding: 8px 24px; border-radius: 6px; font-size: 22px;
                }
                QPushButton:checked { background: #f59e0b; color: white; border-color: #f59e0b; }
            """)
        self.yes_btn.clicked.connect(self._on_yes_clicked)
        self.no_btn.clicked.connect(self._on_no_clicked)
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        input_layout.addLayout(btn_row)

        # Details container (shown when Yes is selected)
        self.details_container = QWidget()
        details_layout = QVBoxLayout(self.details_container)
        details_layout.setContentsMargins(0, 0, 0, 0)
        details_layout.setSpacing(8)

        self.details_edit = QTextEdit()
        self.details_edit.setPlaceholderText("Details of employment opportunities...")
        self.details_edit.setMinimumHeight(100)
        self.details_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.details_edit.setStyleSheet("""
            QTextEdit {
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                background: white;
                font-size: 17px;
            }
            QTextEdit:focus { border-color: #f59e0b; }
        """)
        details_layout.addWidget(self.details_edit)

        # Drag handle for resizing the text box
        drag_handle = QFrame()
        drag_handle.setFixedHeight(10)
        drag_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        drag_handle.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25),
                    stop:0.6 rgba(0,0,0,0.15), stop:1 transparent);
                border: none;
                border-radius: 3px;
                margin: 0 30%;
            }
        """)
        drag_handle._drag_start_y = None
        drag_handle._initial_height = None
        _edit_ref = self.details_edit

        def _handle_press(event):
            drag_handle._drag_start_y = event.globalPosition().y()
            drag_handle._initial_height = _edit_ref.height()
        def _handle_move(event):
            if drag_handle._drag_start_y is not None:
                delta = int(event.globalPosition().y() - drag_handle._drag_start_y)
                new_h = max(80, drag_handle._initial_height + delta)
                _edit_ref.setMinimumHeight(new_h)
                _edit_ref.setMaximumHeight(new_h)
        def _handle_release(event):
            if drag_handle._drag_start_y is not None:
                _edit_ref.setMaximumHeight(16777215)  # Reset max
                drag_handle._drag_start_y = None

        drag_handle.mousePressEvent = _handle_press
        drag_handle.mouseMoveEvent = _handle_move
        drag_handle.mouseReleaseEvent = _handle_release
        details_layout.addWidget(drag_handle)

        self.details_container.hide()
        input_layout.addWidget(self.details_container)

        # Send button
        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(self._send)
        input_layout.addWidget(send_btn)

        main_layout.addWidget(input_frame)

        # IMPORTED DATA SECTION
        self.extracted_section = CollapsibleSection("Imported Data", start_collapsed=True)
        self.extracted_section.set_content_height(200)
        self.extracted_section._min_height = 60
        self.extracted_section._max_height = 2000
        self.extracted_section.set_header_style("QFrame { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-radius: 6px 6px 0 0; }")
        self.extracted_section.set_title_style("QLabel { font-size: 18px; font-weight: 600; color: #806000; background: transparent; border: none; }")

        extracted_content = QWidget()
        extracted_content.setStyleSheet("QWidget { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-top: none; border-radius: 0 0 12px 12px; }")

        extracted_layout = QVBoxLayout(extracted_content)
        extracted_layout.setContentsMargins(12, 10, 12, 10)
        extracted_layout.setSpacing(6)

        extracted_scroll = QScrollArea()
        extracted_scroll.setWidgetResizable(True)
        extracted_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        extracted_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        extracted_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        extracted_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        self.extracted_container = QWidget()
        self.extracted_container.setStyleSheet("background: transparent;")
        self.extracted_checkboxes_layout = QVBoxLayout(self.extracted_container)
        self.extracted_checkboxes_layout.setContentsMargins(2, 2, 2, 2)
        self.extracted_checkboxes_layout.setSpacing(12)

        extracted_scroll.setWidget(self.extracted_container)
        extracted_layout.addWidget(extracted_scroll)

        self.extracted_section.set_content(extracted_content)
        self.extracted_section.setVisible(False)
        main_layout.addWidget(self.extracted_section)

        main_layout.addStretch()
        main_scroll.setWidget(main_container)
        root.addWidget(main_scroll)

    def _send_to_card(self):
        """Send checked items directly to card (replaces card content)."""
        parts = [cb.property("full_text") for cb in self._extracted_checkboxes if cb.isChecked()]
        self.sent.emit("\n\n".join(parts))


    def _on_yes_clicked(self):
        """Handle Yes button click."""
        self.no_btn.setChecked(False)
        self.details_container.show()

    def _on_no_clicked(self):
        """Handle No button click."""
        self.yes_btn.setChecked(False)
        self.details_container.hide()

    def _send(self):
        """Send employment data to card."""
        if self.no_btn.isChecked():
            self.sent.emit("There are no opportunities for employment if discharged.")
        elif self.yes_btn.isChecked():
            details = self.details_edit.toPlainText().strip()
            if details:
                self.sent.emit(f"Yes. {details}")
            else:
                self.sent.emit("Yes. Employment opportunities are available if discharged.")

    def set_entries(self, entries: list, date_range_info: str = ""):
        self._entries = entries

        for cb in self._extracted_checkboxes:
            cb.deleteLater()
        self._extracted_checkboxes.clear()

        while self.extracted_checkboxes_layout.count():
            item = self.extracted_checkboxes_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if isinstance(entries, str):
            entries = [{"date": None, "text": p.strip()} for p in entries.split("\n\n") if p.strip()] if entries.strip() else []

        if entries:
            def get_sort_date(item):
                dt = item.get("date")
                return dt.strftime("%Y-%m-%d") if dt and hasattr(dt, "strftime") else str(dt) if dt else ""

            sorted_items = sorted(entries, key=get_sort_date, reverse=True)

            for item in sorted_items:
                dt = item.get("date")
                text = (item.get("text", "") or item.get("content", "")).strip()
                if not text:
                    continue

                date_str = dt.strftime("%d %b %Y") if dt and hasattr(dt, "strftime") else str(dt) if dt else "No date"

                entry_frame = QFrame()
                entry_frame.setObjectName("entryFrame")
                entry_frame.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                entry_frame.setStyleSheet("QFrame#entryFrame { background: rgba(255,255,255,0.95); border: 1px solid rgba(180,150,50,0.4); border-radius: 8px; padding: 4px; }")
                entry_layout = QVBoxLayout(entry_frame)
                entry_layout.setContentsMargins(10, 8, 10, 8)
                entry_layout.setSpacing(6)
                entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

                header_row = QHBoxLayout()
                header_row.setSpacing(8)

                # Toggle button on LEFT with arrow
                toggle_btn = QPushButton("▸")
                toggle_btn.setFixedSize(22, 22)
                toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                toggle_btn.setStyleSheet("QPushButton { background: rgba(180,150,50,0.2); border: none; border-radius: 4px; font-size: 17px; font-weight: bold; color: #806000; } QPushButton:hover { background: rgba(180,150,50,0.35); }")
                header_row.addWidget(toggle_btn)

                date_label = QLabel(f"{date_str}")
                date_label.setStyleSheet("QLabel { font-size: 17px; font-weight: 600; color: #806000; background: transparent; border: none; }")
                date_label.setCursor(Qt.CursorShape.PointingHandCursor)
                header_row.addWidget(date_label)
                header_row.addStretch()

                # Checkbox on RIGHT
                cb = QCheckBox()
                cb.setProperty("full_text", text)
                cb.setFixedSize(18, 18)
                cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.stateChanged.connect(self._send_to_card)
                header_row.addWidget(cb)

                entry_layout.addLayout(header_row)

                body_text = QTextEdit()
                body_text.setPlainText(text)
                body_text.setReadOnly(True)
                body_text.setFrameShape(QFrame.Shape.NoFrame)
                body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                body_text.setStyleSheet("QTextEdit { font-size: 17px; color: #333; background: rgba(255,248,220,0.5); border: none; padding: 8px; border-radius: 6px; }")
                body_text.setMinimumHeight(60)
                body_text.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
                body_text.setVisible(False)
                entry_layout.addWidget(body_text)

                drag_bar = QFrame()
                drag_bar.setFixedHeight(8)
                drag_bar.setCursor(Qt.CursorShape.SizeVerCursor)
                drag_bar.setStyleSheet("""
                    QFrame {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.3), stop:1 rgba(180,150,50,0.1));
                        border-radius: 2px; margin: 2px 40px;
                    }
                    QFrame:hover {
                        background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                            stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.5), stop:1 rgba(180,150,50,0.2));
                    }
                """)
                drag_bar.setVisible(False)
                drag_bar._drag_y = None
                drag_bar._init_h = None
                def _make_drag_handlers(handle, text_widget):
                    def press(ev):
                        handle._drag_y = ev.globalPosition().y()
                        handle._init_h = text_widget.height()
                    def move(ev):
                        if handle._drag_y is not None:
                            delta = int(ev.globalPosition().y() - handle._drag_y)
                            new_h = max(60, handle._init_h + delta)
                            text_widget.setMinimumHeight(new_h)
                            text_widget.setMaximumHeight(new_h)
                    def release(ev):
                        if handle._drag_y is not None:
                            text_widget.setMaximumHeight(16777215)
                            handle._drag_y = None
                    return press, move, release
                dp, dm, dr = _make_drag_handlers(drag_bar, body_text)
                drag_bar.mousePressEvent = dp
                drag_bar.mouseMoveEvent = dm
                drag_bar.mouseReleaseEvent = dr
                entry_layout.addWidget(drag_bar)

                def make_toggle(btn, body, frame, popup_self, bar):
                    def toggle():
                        is_visible = body.isVisible()
                        body.setVisible(not is_visible)
                        bar.setVisible(not is_visible)
                        btn.setText("▾" if not is_visible else "▸")
                        frame.updateGeometry()
                        popup_self.extracted_container.updateGeometry()
                    return toggle

                toggle_fn = make_toggle(toggle_btn, body_text, entry_frame, self, drag_bar)
                toggle_btn.clicked.connect(toggle_fn)
                date_label.mousePressEvent = lambda e, fn=toggle_fn: fn()

                self.extracted_checkboxes_layout.addWidget(entry_frame)
                self._extracted_checkboxes.append(cb)

            self.extracted_checkboxes_layout.addStretch()
            self.extracted_section.setVisible(True)
            # Keep collapsed on open - user can expand manually
        else:
            self.extracted_section.setVisible(False)


class PreviousCommunityPopup(QWidget):
    """Popup for previous response to community support."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("What was the patient's previous response to community support or Section 117 aftercare?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Describe previous response to community support...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class CarePathwayPopup(QWidget):
    """Popup for care pathway and Section 117 aftercare."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("What care pathway and Section 117 after-care will be made available to the patient?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Describe care pathway and Section 117 aftercare available...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class CarePlanPopup(QWidget):
    """Popup for proposed care plan."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Give details of the proposed care plan:")
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Describe the proposed care plan...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class CarePlanAdequacyPopup(QWidget):
    """Popup for care plan adequacy."""
    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("How adequate or effective is the proposed care plan likely to be?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Comment on adequacy and effectiveness of care plan...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class CarePlanFundingPopup(QWidget):
    """Popup for care plan funding issues."""
    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Are there any issues as to funding the proposed care plan?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes")
        self.no_btn = QPushButton("No")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6; border: 1px solid #d1d5db;
                    padding: 8px 24px; border-radius: 6px; font-size: 22px;
                }
                QPushButton:checked { background: #f59e0b; color: white; border-color: #f59e0b; }
            """)
        self.yes_btn.clicked.connect(lambda: self.no_btn.setChecked(False))
        self.no_btn.clicked.connect(lambda: self.yes_btn.setChecked(False))
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        self.details_edit = QTextEdit()
        self.details_edit.setPlaceholderText("If yes, by what date will those issues be resolved?")
        self.details_edit.setMinimumHeight(80)
        self.details_edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.details_edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.details_edit)

        # Drag handle for resizing the text box
        _det_handle = QFrame()
        _det_handle.setFixedHeight(10)
        _det_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _det_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _det_handle._drag_y = None
        _det_handle._init_h = None
        _det_handle._init_p_h = None
        _de = self.details_edit
        _popup_ref14 = self
        def _dp(ev, h=_det_handle, e=_de, p=_popup_ref14):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _dm(ev, h=_det_handle, e=_de, p=_popup_ref14):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _dr(ev, h=_det_handle, e=_de, p=_popup_ref14):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _det_handle.mousePressEvent = _dp
        _det_handle.mouseMoveEvent = _dm
        _det_handle.mouseReleaseEvent = _dr
        layout.addWidget(_det_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(self._send)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)

    def _send(self):
        if self.no_btn.isChecked():
            self.sent.emit("There are no issues as to funding the proposed care plan.")
        elif self.yes_btn.isChecked():
            details = self.details_edit.toPlainText().strip()
            if details:
                self.sent.emit(f"Yes. {details}")
            else:
                self.sent.emit("Yes, there are funding issues (details to be provided).")


class PatientViewsPopup(QWidget):
    """Popup for patient's views, wishes, beliefs, opinions, hopes and concerns - sends to card on click."""
    sent = Signal(str)

    DISCHARGE_OPTIONS = [
        ("", "-- Select --"),
        ("keen", "is keen to be discharged"),
        ("moving_towards", "is moving towards discharge"),
        ("ambivalent", "is ambivalent about discharge"),
        ("unsure", "is unsure about discharge"),
    ]

    def __init__(self, parent=None, gender=None):
        super().__init__(parent)
        self._gender = gender
        self._set_pronoun(gender)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _set_pronoun(self, gender):
        """Set pronoun based on gender."""
        if gender == "Male":
            self._pronoun = "He"
        elif gender == "Female":
            self._pronoun = "She"
        else:
            self._pronoun = "They"

    def _setup_ui(self):
        """Set up the UI components."""
        layout = self.layout()

        # ========== INPUT PANEL ==========
        input_frame = QFrame()
        input_frame.setStyleSheet("""
            QFrame {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
            }
        """)
        input_layout = QVBoxLayout(input_frame)
        input_layout.setContentsMargins(12, 12, 12, 12)
        input_layout.setSpacing(12)

        # Scroll area for inputs
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(16)

        # --- Discharge Views ---
        discharge_group = QFrame()
        discharge_group.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; padding: 8px; }")
        discharge_layout = QVBoxLayout(discharge_group)
        discharge_layout.setContentsMargins(12, 8, 12, 8)
        discharge_layout.setSpacing(6)

        discharge_label = QLabel("Patient's views on discharge:")
        discharge_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        discharge_layout.addWidget(discharge_label)

        self.discharge_combo = QComboBox()
        for key, text in self.DISCHARGE_OPTIONS:
            self.discharge_combo.addItem(text, key)
        self.discharge_combo.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                background: white;
                font-size: 22px;
            }
            QComboBox:hover { border-color: #f59e0b; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        self.discharge_combo.currentIndexChanged.connect(self._update_preview)
        discharge_layout.addWidget(self.discharge_combo)
        scroll_layout.addWidget(discharge_group)

        # --- Wishes and Hopes ---
        wishes_group = QFrame()
        wishes_group.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; padding: 8px; }")
        wishes_layout = QVBoxLayout(wishes_group)
        wishes_layout.setContentsMargins(12, 8, 12, 8)
        wishes_layout.setSpacing(6)

        wishes_label = QLabel("Patient's wishes and hopes:")
        wishes_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        wishes_layout.addWidget(wishes_label)

        self.wishes_edit = QTextEdit()
        self.wishes_edit.setPlaceholderText("Enter patient's wishes and hopes...")
        self.wishes_edit.setMaximumHeight(80)
        self.wishes_edit.setStyleSheet("QTextEdit { font-size: 22px; }")
        self.wishes_edit.textChanged.connect(self._update_preview)
        wishes_layout.addWidget(self.wishes_edit)
        scroll_layout.addWidget(wishes_group)

        # --- Beliefs and Opinions ---
        beliefs_group = QFrame()
        beliefs_group.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; padding: 8px; }")
        beliefs_layout = QVBoxLayout(beliefs_group)
        beliefs_layout.setContentsMargins(12, 8, 12, 8)
        beliefs_layout.setSpacing(6)

        beliefs_label = QLabel("Patient's beliefs and opinions:")
        beliefs_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        beliefs_layout.addWidget(beliefs_label)

        self.beliefs_edit = QTextEdit()
        self.beliefs_edit.setPlaceholderText("Enter patient's beliefs and opinions...")
        self.beliefs_edit.setMaximumHeight(80)
        self.beliefs_edit.setStyleSheet("QTextEdit { font-size: 22px; }")
        self.beliefs_edit.textChanged.connect(self._update_preview)
        beliefs_layout.addWidget(self.beliefs_edit)
        scroll_layout.addWidget(beliefs_group)

        # --- Concerns ---
        concerns_group = QFrame()
        concerns_group.setStyleSheet("QFrame { background: white; border: 1px solid #e5e7eb; border-radius: 6px; padding: 8px; }")
        concerns_layout = QVBoxLayout(concerns_group)
        concerns_layout.setContentsMargins(12, 8, 12, 8)
        concerns_layout.setSpacing(6)

        concerns_label = QLabel("Patient's concerns:")
        concerns_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        concerns_layout.addWidget(concerns_label)

        self.concerns_edit = QTextEdit()
        self.concerns_edit.setPlaceholderText("Enter patient's concerns...")
        self.concerns_edit.setMaximumHeight(80)
        self.concerns_edit.textChanged.connect(self._update_preview)
        concerns_layout.addWidget(self.concerns_edit)
        scroll_layout.addWidget(concerns_group)

        scroll_layout.addStretch()
        scroll.setWidget(scroll_content)
        input_layout.addWidget(scroll)

        # Send button
        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(self._send)
        input_layout.addWidget(send_btn)

        layout.addWidget(input_frame)

    def _generate_text(self):
        """Generate text based on all inputs."""
        parts = []
        pronoun = self._pronoun

        # Discharge view
        discharge_key = self.discharge_combo.currentData()
        if discharge_key:
            discharge_text = self.discharge_combo.currentText()
            parts.append(f"{pronoun} {discharge_text}.")

        # Wishes and hopes
        wishes = self.wishes_edit.toPlainText().strip()
        if wishes:
            parts.append(f"\n{pronoun} wishes and hopes: {wishes}")

        # Beliefs and opinions
        beliefs = self.beliefs_edit.toPlainText().strip()
        if beliefs:
            parts.append(f"\n{pronoun} believes: {beliefs}")

        # Concerns
        concerns = self.concerns_edit.toPlainText().strip()
        if concerns:
            parts.append(f"\n{pronoun} is concerned about: {concerns}")

        return ' '.join(parts) if parts else ""

    def _update_preview(self):
        """Send text to card immediately."""
        text = self._generate_text()
        if text:
            self.sent.emit(text)

    def _send(self):
        text = self._generate_text()
        if text:
            self.sent.emit(text)

    def update_gender(self, gender: str):
        """Update gender and refresh preview."""
        self._gender = gender
        self._set_pronoun(gender)
        self._update_preview()


class NearestRelativePopup(QWidget):
    """Popup for nearest relative views."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Other than in restricted cases, what are the views of the patient's Nearest Relative?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Describe the Nearest Relative's views...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class NRInappropriatePopup(QWidget):
    """Popup for when inappropriate to consult Nearest Relative."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("If (having consulted the patient) it was considered\ninappropriate or impractical to consult the Nearest Relative,\nwhat were the reasons for this and what attempts\nhave been made to rectify matters?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Explain reasons and attempts made...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class CarerViewsPopup(QWidget):
    """Popup for views of other person taking lead role in care."""
    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("What are the views of any other person who takes\na lead role in the care and support of the patient\nbut who is not professionally involved?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        self.edit = QTextEdit()
        self.edit.setPlaceholderText("Describe the views of non-professional carers...")
        self.edit.setMinimumHeight(150)
        self.edit.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.edit.setStyleSheet("font-size: 17px;")
        layout.addWidget(self.edit)

        # Drag handle for resizing the text box
        _edit_handle = QFrame()
        _edit_handle.setFixedHeight(10)
        _edit_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        _edit_handle.setStyleSheet("QFrame { background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 transparent, stop:0.4 rgba(0,0,0,0.15), stop:0.5 rgba(0,0,0,0.25), stop:0.6 rgba(0,0,0,0.15), stop:1 transparent); border: none; border-radius: 3px; margin: 0 30%; }")
        _edit_handle._drag_y = None
        _edit_handle._init_h = None
        _edit_handle._init_p_h = None
        _ed = self.edit
        _popup_ref = self
        def _ep(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            h._drag_y = ev.globalPosition().y(); h._init_h = e.height()
            h._init_p_h = p.height()
        def _em(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                new_h = max(80, h._init_h + delta)
                e.setMinimumHeight(new_h)
                e.setMaximumHeight(new_h)
                new_p_h = max(200, h._init_p_h + delta)
                p.setMinimumHeight(new_p_h)
                p.setMaximumHeight(new_p_h)
        def _er(ev, h=_edit_handle, e=_ed, p=_popup_ref):
            e.setMaximumHeight(16777215); p.setMinimumHeight(0); p.setMaximumHeight(16777215); h._drag_y = None
        _edit_handle.mousePressEvent = _ep
        _edit_handle.mouseMoveEvent = _em
        _edit_handle.mouseReleaseEvent = _er
        layout.addWidget(_edit_handle)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(lambda: self.sent.emit(self.edit.toPlainText().strip()) if self.edit.toPlainText().strip() else None)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)


class MAPPAPopup(QWidget):
    """Popup for MAPPA involvement - sends to card on click."""
    sent = Signal(str)

    def __init__(self, parent=None, gender=None):
        super().__init__(parent)
        self._gender = gender
        self._set_pronouns(gender)

        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # ========== INPUT PANEL ==========
        input_widget = QWidget()
        input_layout = QVBoxLayout(input_widget)
        input_layout.setContentsMargins(16, 16, 16, 16)
        input_layout.setSpacing(12)

        # First question - Is patient known to MAPPA?
        label = QLabel("Is the patient known to any MAPPA meeting or agency?")
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        input_layout.addWidget(label)

        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes")
        self.no_btn = QPushButton("No")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6; border: 1px solid #d1d5db;
                    padding: 8px 24px; border-radius: 6px; font-size: 22px;
                }
                QPushButton:checked { background: #f59e0b; color: white; border-color: #f59e0b; }
            """)
        self.yes_btn.clicked.connect(self._on_yes_clicked)
        self.no_btn.clicked.connect(self._on_no_clicked)
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        input_layout.addLayout(btn_row)

        # Container for details (shown when Yes is selected)
        self.details_container = QWidget()
        details_layout = QVBoxLayout(self.details_container)
        details_layout.setContentsMargins(0, 0, 0, 0)
        details_layout.setSpacing(12)

        # MAPPA Management Level dropdown
        level_label = QLabel("MAPPA Management Level:")
        level_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        details_layout.addWidget(level_label)

        self.level_combo = QComboBox()
        self.level_combo.addItems([
            "Select level...",
            "1 - Ordinary agency",
            "2 - Active multi-agency",
            "3 - Enhanced multi-agency"
        ])
        self.level_combo.setStyleSheet("""
            QComboBox {
                padding: 8px 12px; border: 1px solid #d1d5db;
                border-radius: 6px; background: white; font-size: 22px;
            }
            QComboBox:hover { border-color: #f59e0b; }
            QComboBox::drop-down { border: none; padding-right: 8px; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        self.level_combo.currentIndexChanged.connect(self._update_preview)
        details_layout.addWidget(self.level_combo)

        # MAPPA Offender Categories dropdown
        category_label = QLabel("MAPPA Offender Category:")
        category_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        details_layout.addWidget(category_label)

        self.category_combo = QComboBox()
        self.category_combo.addItems([
            "Select category...",
            "1 - Registered sex offenders",
            "2 - Violent/serious sexual",
            "3 - Other dangerous",
            "4 - Terrorism-related"
        ])
        self.category_combo.setStyleSheet("""
            QComboBox {
                padding: 8px 12px; border: 1px solid #d1d5db;
                border-radius: 6px; background: white; font-size: 22px;
            }
            QComboBox:hover { border-color: #f59e0b; }
            QComboBox::drop-down { border: none; padding-right: 8px; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        self.category_combo.currentIndexChanged.connect(self._update_preview)
        details_layout.addWidget(self.category_combo)

        # Chair of MAPPA meeting
        chair_label = QLabel("Name of Chair of MAPPA meeting:")
        chair_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        details_layout.addWidget(chair_label)

        self.chair_edit = QLineEdit()
        self.chair_edit.setPlaceholderText("Enter chair name...")
        self.chair_edit.setStyleSheet("""
            QLineEdit {
                padding: 8px 12px; border: 1px solid #d1d5db;
                border-radius: 6px; background: white; font-size: 22px;
            }
            QLineEdit:focus { border-color: #f59e0b; }
        """)
        self.chair_edit.textChanged.connect(self._update_preview)
        details_layout.addWidget(self.chair_edit)

        # Lead agency representative
        lead_label = QLabel("Name of representative of lead agency:")
        lead_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        details_layout.addWidget(lead_label)

        self.lead_edit = QLineEdit()
        self.lead_edit.setPlaceholderText("Enter lead name...")
        self.lead_edit.setStyleSheet("""
            QLineEdit {
                padding: 8px 12px; border: 1px solid #d1d5db;
                border-radius: 6px; background: white; font-size: 22px;
            }
            QLineEdit:focus { border-color: #f59e0b; }
        """)
        self.lead_edit.textChanged.connect(self._update_preview)
        details_layout.addWidget(self.lead_edit)

        self.details_container.hide()  # Hidden by default
        input_layout.addWidget(self.details_container)

        input_layout.addStretch()

        # Send to Report button at bottom
        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(self._send)
        input_layout.addWidget(send_btn)

        main_layout.addWidget(input_widget)

        add_lock_to_popup(self, show_button=False)

    def _set_pronouns(self, gender):
        """Set pronouns based on gender."""
        if gender == "Male":
            self._pronoun = "He"
            self._pronoun_lower = "he"
            self._possessive = "His"
        elif gender == "Female":
            self._pronoun = "She"
            self._pronoun_lower = "she"
            self._possessive = "Her"
        else:
            self._pronoun = "They"
            self._pronoun_lower = "they"
            self._possessive = "Their"

    def update_gender(self, gender: str):
        """Update gender and refresh card."""
        self._gender = gender
        self._set_pronouns(gender)
        self._update_preview()

    def _on_yes_clicked(self):
        self.no_btn.setChecked(False)
        self.details_container.show()
        self._update_preview()

    def _on_no_clicked(self):
        self.yes_btn.setChecked(False)
        self.details_container.hide()
        self._update_preview()

    def _update_preview(self):
        """Send text directly to card."""
        text = self._generate_text()
        if text:
            self.sent.emit(text)

    def _generate_text(self):
        """Generate the MAPPA text based on selections."""
        if self.no_btn.isChecked():
            return f"{self._pronoun} is not known to any MAPPA meeting or agency."

        if not self.yes_btn.isChecked():
            return ""

        parts = []
        parts.append(f"{self._pronoun} is known to MAPPA.")

        level_idx = self.level_combo.currentIndex()
        if level_idx > 0:
            level_text = self.level_combo.currentText()
            parts.append(f"{self._possessive} case is managed at {level_text}.")

        category_idx = self.category_combo.currentIndex()
        if category_idx > 0:
            category_text = self.category_combo.currentText()
            parts.append(f"{self._pronoun} falls under MAPPA Category {category_text}.")

        chair = self.chair_edit.text().strip()
        if chair:
            parts.append(f"The Chair of the MAPPA meeting is {chair}.")

        lead = self.lead_edit.text().strip()
        if lead:
            parts.append(f"The lead agency representative is {lead}.")

        return ' '.join(parts)

    def _send(self):
        text = self._generate_text()
        if text:
            self.sent.emit(text)
        elif self.yes_btn.isChecked():
            self.sent.emit(f"{self._pronoun} is known to MAPPA (details to be provided).")
        elif self.no_btn.isChecked():
            self.sent.emit(f"{self._pronoun} is not known to any MAPPA meeting or agency.")


class OtherInfoPopup(QWidget):
    """Popup for other relevant information."""
    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Is there any other relevant information that the tribunal should know?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes")
        self.no_btn = QPushButton("No")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6; border: 1px solid #d1d5db;
                    padding: 8px 24px; border-radius: 6px; font-size: 22px;
                }
                QPushButton:checked { background: #f59e0b; color: white; border-color: #f59e0b; }
            """)
        self.yes_btn.clicked.connect(lambda: self.no_btn.setChecked(False))
        self.no_btn.clicked.connect(lambda: self.yes_btn.setChecked(False))
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        self.details_edit = QTextEdit()
        self.details_edit.setStyleSheet("font-size: 22px;")
        self.details_edit.setPlaceholderText("Provide other relevant information...")
        self.details_edit.setMinimumHeight(120)
        layout.addWidget(self.details_edit)

        layout.addStretch()

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b; color: white; border: none;
                padding: 10px 20px; border-radius: 8px; font-weight: 600; font-size: 22px;
            }
            QPushButton:hover { background: #d97706; }
        """)
        send_btn.clicked.connect(self._send)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)

    def _send(self):
        if self.no_btn.isChecked():
            self.sent.emit("There is no other relevant information for the tribunal.")
        elif self.yes_btn.isChecked():
            details = self.details_edit.toPlainText().strip()
            if details:
                self.sent.emit(details)
