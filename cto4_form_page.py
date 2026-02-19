# ================================================================
#  CTO4 FORM PAGE — Record of Patient's Detention After Recall
#  Mental Health Act 1983 - Form CTO4 Regulation 6(3)(d)
#  Section 17E — Community treatment order: record of detention
#  CARD/POPUP LAYOUT
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QPushButton, QFileDialog, QMessageBox, QToolButton,
    QStackedWidget, QSplitter
)
from shared_widgets import create_zoom_row
from background_history_popup import ResizableSection
from utils.resource_path import resource_path
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# CTO4 CARD WIDGET
# ================================================================
class CTO4CardWidget(QFrame):
    """Clickable card with editable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("cto4Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)

        self.setStyleSheet("""
            QFrame#cto4Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#cto4Card:hover {
                border-color: #7c3aed;
                background: #faf5ff;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        self.setMinimumHeight(80)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(6)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 19px; font-weight: 700; color: #7c3aed;")
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 19px;
                color: #374151;
                background: transparent;
                border: none;
                padding: 0;
            }
        """)
        self.content.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        layout.addWidget(self.content, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=17)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_active(self, active: bool):
        """Set active state - shows persistent highlight."""
        if active:
            self.setStyleSheet("""
                QFrame#cto4Card {
                    background: #faf5ff;
                    border: 2px solid #7c3aed;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#cto4Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#cto4Card:hover {
                    border-color: #7c3aed;
                    background: #faf5ff;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# CTO4 TOOLBAR
# ================================================================
class CTO4Toolbar(QWidget):
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            CTO4Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN CTO4 FORM PAGE - Card/Popup Layout
# ================================================================
class CTO4FormPage(QWidget):
    """Page for completing MHA Form CTO4 - Record of Detention After Recall."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}

        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {"full_name": details[1] or "", "email": details[7] or ""}

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.signatory_name.setText(self._my_details["full_name"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(50)
        header.setStyleSheet("background: #7c3aed; border-bottom: 1px solid #6d28d9;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 5px;
                font-size: 18px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO4 — Record of Detention After Recall")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - cards on left, popup on right with splitter
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(12, 12, 12, 12)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #7c3aed, stop:0.7 #7c3aed, stop:1 transparent);
            }
        """)

        # Left: Cards column (scrollable)
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(0, 0, 8, 0)
        self.cards_layout.setSpacing(8)

        # Create cards
        self._create_details_card()
        self._create_detention_card()
        self._create_signature_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        self.main_splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 10px;
            }
        """)

        # Create popup panels
        self._create_details_popup()
        self._create_detention_popup()
        self._create_signature_popup()

        # Initialize cards with default date values
        self._update_detention_card()
        self._update_signature_card()

        self.main_splitter.addWidget(self.popup_stack)
        self.main_splitter.setSizes([280, 600])
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        # Show first popup by default
        self._on_card_clicked("details")

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"details": 0, "detention": 1, "signature": 2}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            for k, card in self.cards.items():
                card.set_active(k == key)

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_details_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 350
        card = CTO4CardWidget("Patient & Hospital", "details")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["details"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_detention_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 350
        card = CTO4CardWidget("Detention Details", "detention")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["detention"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signature_card(self):
        section = ResizableSection()
        section.set_content_height(150)
        section._min_height = 100
        section._max_height = 300
        card = CTO4CardWidget("Signatory", "signature")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signature"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_details_popup(self):
        """Popup for patient and hospital details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Patient & Hospital Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #7c3aed;")
        popup_layout.addWidget(header)

        # Form
        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        # Patient section
        patient_header = QLabel("Patient")
        patient_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 4px;")
        form_layout.addWidget(patient_header)

        self.patient_name = self._create_line_edit("Patient full name")
        self.patient_name.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.patient_name)

        self.patient_address = self._create_line_edit("Patient address")
        self.patient_address.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.patient_address)

        # Hospital section
        hosp_header = QLabel("Hospital")
        hosp_header.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(hosp_header)

        self.hospital_name = self._create_line_edit("Hospital name")
        self.hospital_name.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.hospital_name)

        self.hospital_address = self._create_line_edit("Hospital address")
        self.hospital_address.textChanged.connect(self._update_details_card)
        form_layout.addWidget(self.hospital_address)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        self.popup_stack.addWidget(popup)

    def _create_detention_popup(self):
        """Popup for detention date and time."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Detention Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #059669;")
        popup_layout.addWidget(header)

        # Info text
        info = QLabel("The patient arrived at the hospital in pursuance of a notice recalling them under section 17E:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 18px; color: #374151; padding: 8px; background: #f0fdf4; border-radius: 6px;")
        popup_layout.addWidget(info)

        # Form
        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        # Date row
        date_row = QHBoxLayout()
        date_row.setSpacing(8)
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        date_row.addWidget(date_lbl)
        self.detention_date = self._create_date_edit()
        self.detention_date.setFixedWidth(140)
        self.detention_date.dateChanged.connect(self._update_detention_card)
        date_row.addWidget(self.detention_date)
        date_row.addStretch()
        form_layout.addLayout(date_row)

        # Time row
        time_row = QHBoxLayout()
        time_row.setSpacing(8)
        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        time_row.addWidget(time_lbl)
        self.detention_time = self._create_time_edit()
        self.detention_time.timeChanged.connect(self._update_detention_card)
        self.detention_time.setFixedWidth(100)
        time_row.addWidget(self.detention_time)
        time_row.addStretch()
        form_layout.addLayout(time_row)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        self.popup_stack.addWidget(popup)

    def _create_signature_popup(self):
        """Popup for signatory details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Signatory Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #dc2626;")
        popup_layout.addWidget(header)

        # Form
        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        name_lbl = QLabel("Name (on behalf of hospital managers):")
        name_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        form_layout.addWidget(name_lbl)

        self.signatory_name = self._create_line_edit("Signatory full name")
        self.signatory_name.textChanged.connect(self._update_signature_card)
        form_layout.addWidget(self.signatory_name)

        # Date row
        date_row = QHBoxLayout()
        date_row.setSpacing(8)
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        date_row.addWidget(date_lbl)
        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        self.sig_date.dateChanged.connect(self._update_signature_card)
        date_row.addWidget(self.sig_date)
        date_row.addStretch()
        form_layout.addLayout(date_row)

        # Time row
        time_row = QHBoxLayout()
        time_row.setSpacing(8)
        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        time_row.addWidget(time_lbl)
        self.sig_time = self._create_time_edit()
        self.sig_time.setFixedWidth(100)
        self.sig_time.timeChanged.connect(self._update_signature_card)
        time_row.addWidget(self.sig_time)
        time_row.addStretch()
        form_layout.addLayout(time_row)

        notice = QLabel("This record must be made by the managers of the hospital to which the patient has been recalled.")
        notice.setWordWrap(True)
        notice.setStyleSheet("font-size: 17px; color: #6b7280; font-style: italic; padding: 8px; background: #f3f4f6; border-radius: 6px; margin-top: 12px;")
        form_layout.addWidget(notice)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_details_card(self):
        parts = []
        if self.patient_name.text().strip():
            parts.append(self.patient_name.text().strip())
        if self.patient_address.text().strip():
            parts.append(self.patient_address.text().strip()[:40] + "...")
        if self.hospital_name.text().strip():
            parts.append(self.hospital_name.text().strip())
        if self.hospital_address.text().strip():
            parts.append(self.hospital_address.text().strip()[:40] + "...")
        self.cards["details"].set_content_text("\n".join(parts) if parts else "Click to enter details")

    def _update_detention_card(self):
        parts = []
        if self.detention_date.date().isValid():
            parts.append(self.detention_date.date().toString('dd MMM yyyy'))
        if self.detention_time.time().isValid():
            parts.append(self.detention_time.time().toString('HH:mm'))
        self.cards["detention"].set_content_text("\n".join(parts) if parts else "Click to enter details")

    def _update_signature_card(self):
        parts = []
        if self.signatory_name.text().strip():
            parts.append(self.signatory_name.text().strip())
        if self.sig_date.date().isValid():
            parts.append(self.sig_date.date().toString('dd MMM yyyy'))
        if self.sig_time.time().isValid():
            parts.append(self.sig_time.time().toString('HH:mm'))
        self.cards["signature"].set_content_text("\n".join(parts) if parts else "Click to enter details")

    # ----------------------------------------------------------------
    # HELPERS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("QLineEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 8px; font-size: 18px; } QLineEdit:focus { border-color: #7c3aed; }")
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 6px; font-size: 18px; }")
        return date_edit

    def _create_time_edit(self) -> QTimeEdit:
        time_edit = QTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("QTimeEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 6px; font-size: 18px; }")
        return time_edit

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.patient_name.clear()
            self.patient_address.clear()
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.detention_date.setDate(QDate.currentDate())
            self.detention_time.setTime(QTime.currentTime())
            self.signatory_name.clear()
            self.sig_date.setDate(QDate.currentDate())
            self.sig_time.setTime(QTime.currentTime())
            for card in self.cards.values():
                card.set_content_text("")
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form CTO4",
            f"Form_CTO4_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_CTO4_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO4 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            from docx.shared import RGBColor, Pt

            # Gold bracket color (#918C0D) and cream highlight (#FFFED5)
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)
            CREAM_FILL = 'FFFED5'

            # Clean ALL paragraphs - remove permission markers and convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                # Remove permission markers
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                # Convert paragraph-level shading to cream
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), CREAM_FILL)
                # Convert run-level shading to cream
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), CREAM_FILL)

            def set_entry_box(para, content: str):
                """Set entry box with bold gold brackets [content] - only content between brackets is cream."""
                if not content or not content.strip():
                    content = '                                                                   '
                # Clear all runs
                for run in para.runs:
                    run.text = ""
                # Remove extra runs
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)

                # Remove any paragraph-level shading
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)

                # Add opening bracket with gold color and bold (no highlight)
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR

                # Add content with cream highlighting
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                for old_shd in rPr2.findall(qn('w:shd')):
                    rPr2.remove(old_shd)
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)

                # Add closing bracket with gold color and bold (no highlight)
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR

            def set_labeled_entry(para, label: str, content: str, suffix: str = ""):
                """Set labeled entry like 'Date[    ]' - only content between brackets is cream."""
                if not content or not content.strip():
                    content = '                                        '
                # Clear all runs
                for run in para.runs:
                    run.text = ""
                # Remove extra runs
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)

                # Remove any paragraph-level shading
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)

                # Add label (no highlight)
                if para.runs:
                    para.runs[0].text = label
                    label_run = para.runs[0]
                else:
                    label_run = para.add_run(label)
                label_run.font.name = 'Arial'
                label_run.font.size = Pt(12)

                # Add opening bracket with gold color and bold (no highlight)
                bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR

                # Add content with cream highlighting
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                for old_shd in rPr2.findall(qn('w:shd')):
                    rPr2.remove(old_shd)
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)

                # Add closing bracket with gold color and bold (no highlight)
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR

                # Add suffix if any (no highlight)
                if suffix:
                    suffix_run = para.add_run(suffix)
                    suffix_run.font.name = 'Arial'
                    suffix_run.font.size = Pt(12)

            # Build data strings
            patient_text = self.patient_name.text().strip()
            if self.patient_address.text().strip():
                patient_text += ", " + self.patient_address.text().strip()

            hospital_text = self.hospital_name.text().strip()
            if self.hospital_address.text().strip():
                hospital_text += ", " + self.hospital_address.text().strip()

            det_date = self.detention_date.date().toString("dd/MM/yyyy")
            det_time = self.detention_time.time().toString("HH:mm")
            sig_name = self.signatory_name.text().strip()

            # Debug: print paragraph structure
            print("\n=== CTO4 Template Structure ===")
            for i, para in enumerate(doc.paragraphs):
                txt = para.text[:80].replace('\n', ' ') if para.text else "(empty)"
                print(f"DEBUG: Para {i}: {txt}")
            print("=== End Structure ===\n")

            # Track which entry boxes we've filled
            patient_box_filled = False
            hospital_box_filled = False
            date_time_line_filled = False
            signed_filled = False
            print_name_filled = False
            sig_date_filled = False
            sig_time_filled = False

            # Process paragraphs
            for i, para in enumerate(doc.paragraphs):
                text = para.text
                text_stripped = text.strip()

                # Empty entry box - paragraph with just brackets or empty/whitespace
                # These come after instruction text like "[PRINT full name...]"
                if text_stripped.startswith('[') and text_stripped.endswith(']') and len(text_stripped) < 100:
                    inner = text_stripped[1:-1].strip()
                    # If inner is empty or just spaces, this is an entry box
                    if not inner or inner.isspace() or inner == '':
                        if not patient_box_filled:
                            set_entry_box(para, patient_text)
                            patient_box_filled = True
                            print(f"DEBUG: Filled patient box at para {i}")
                            continue
                        elif not hospital_box_filled:
                            set_entry_box(para, hospital_text)
                            hospital_box_filled = True
                            print(f"DEBUG: Filled hospital box at para {i}")
                            continue

                # Also check for completely empty paragraphs or whitespace-only paragraphs
                # that come after instruction paragraphs
                if not text_stripped or text_stripped.isspace():
                    if not patient_box_filled:
                        set_entry_box(para, patient_text)
                        patient_box_filled = True
                        print(f"DEBUG: Filled patient box (empty para) at para {i}")
                        continue
                    elif not hospital_box_filled:
                        set_entry_box(para, hospital_text)
                        hospital_box_filled = True
                        print(f"DEBUG: Filled hospital box (empty para) at para {i}")
                        continue

                # Date Time line - with or without brackets
                if text_stripped.startswith("Date") and "Time" in text and not date_time_line_filled:
                    # Clear all runs and remove extras
                    for run in para.runs:
                        run.text = ""
                    while len(para.runs) > 1:
                        para._element.remove(para.runs[-1]._element)

                    # Remove any paragraph-level shading
                    pPr = para._element.get_or_add_pPr()
                    for old_shd in pPr.findall(qn('w:shd')):
                        pPr.remove(old_shd)

                    # Date label (no highlight)
                    if para.runs:
                        para.runs[0].text = "Date"
                        date_lbl = para.runs[0]
                    else:
                        date_lbl = para.add_run("Date")
                    date_lbl.font.name = 'Arial'
                    date_lbl.font.size = Pt(12)

                    # Date opening bracket (gold, no highlight)
                    open_b1 = para.add_run("[")
                    open_b1.font.name = 'Arial'
                    open_b1.font.size = Pt(12)
                    open_b1.font.bold = True
                    open_b1.font.color.rgb = BRACKET_COLOR

                    # Date content (cream highlight)
                    date_content = para.add_run(det_date if det_date else "                              ")
                    date_content.font.name = 'Arial'
                    date_content.font.size = Pt(12)
                    rPr1c = date_content._element.get_or_add_rPr()
                    shd1c = OxmlElement('w:shd')
                    shd1c.set(qn('w:val'), 'clear')
                    shd1c.set(qn('w:color'), 'auto')
                    shd1c.set(qn('w:fill'), CREAM_FILL)
                    rPr1c.append(shd1c)

                    # Date closing bracket (gold, no highlight)
                    close_b1 = para.add_run("]")
                    close_b1.font.name = 'Arial'
                    close_b1.font.size = Pt(12)
                    close_b1.font.bold = True
                    close_b1.font.color.rgb = BRACKET_COLOR

                    # Time label (no highlight)
                    time_lbl = para.add_run(" Time")
                    time_lbl.font.name = 'Arial'
                    time_lbl.font.size = Pt(12)

                    # Time opening bracket (gold, no highlight)
                    open_b2 = para.add_run("[")
                    open_b2.font.name = 'Arial'
                    open_b2.font.size = Pt(12)
                    open_b2.font.bold = True
                    open_b2.font.color.rgb = BRACKET_COLOR

                    # Time content (cream highlight)
                    time_content = para.add_run(det_time if det_time else "                              ")
                    time_content.font.name = 'Arial'
                    time_content.font.size = Pt(12)
                    rPr2c = time_content._element.get_or_add_rPr()
                    shd2c = OxmlElement('w:shd')
                    shd2c.set(qn('w:val'), 'clear')
                    shd2c.set(qn('w:color'), 'auto')
                    shd2c.set(qn('w:fill'), CREAM_FILL)
                    rPr2c.append(shd2c)

                    # Time closing bracket (gold, no highlight)
                    close_b2 = para.add_run("]")
                    close_b2.font.name = 'Arial'
                    close_b2.font.size = Pt(12)
                    close_b2.font.bold = True
                    close_b2.font.color.rgb = BRACKET_COLOR

                    date_time_line_filled = True
                    print(f"DEBUG: Filled date/time line at para {i}")
                    continue

                # Signed on behalf of... line - with or without brackets
                if text_stripped.startswith("Signed") and not signed_filled:
                    set_labeled_entry(para, "Signed", "", " on behalf of the hospital managers")
                    signed_filled = True
                    print(f"DEBUG: Filled signed line at para {i}")
                    continue

                # PRINT NAME line - with or without brackets
                if "PRINT NAME" in text and not print_name_filled:
                    set_labeled_entry(para, "PRINT NAME", sig_name)
                    print_name_filled = True
                    print(f"DEBUG: Filled print name at para {i}")
                    continue

                # Date line (signatory) - standalone Date without Time
                if text_stripped.startswith("Date") and "Time" not in text and not sig_date_filled:
                    sig_date_str = self.sig_date.date().toString("dd/MM/yyyy")
                    set_labeled_entry(para, "Date", sig_date_str)
                    sig_date_filled = True
                    print(f"DEBUG: Filled sig date at para {i}")
                    continue

                # Time line (signatory) - standalone Time without Date
                if text_stripped.startswith("Time") and "Date" not in text and not sig_time_filled:
                    sig_time_str = self.sig_time.time().toString("HH:mm")
                    set_labeled_entry(para, "Time", sig_time_str)
                    sig_time_filled = True
                    print(f"DEBUG: Filled sig time at para {i}")
                    continue

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO4 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        return {
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "detention_date": self.detention_date.date().toString("yyyy-MM-dd"),
            "detention_time": self.detention_time.time().toString("HH:mm"),
            "signatory_name": self.signatory_name.text(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
            "sig_time": self.sig_time.time().toString("HH:mm"),
            "cards": {key: card.get_content() for key, card in self.cards.items()},
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        if state.get("detention_date"):
            self.detention_date.setDate(QDate.fromString(state["detention_date"], "yyyy-MM-dd"))
        if state.get("detention_time"):
            self.detention_time.setTime(QTime.fromString(state["detention_time"], "HH:mm"))
        self.signatory_name.setText(state.get("signatory_name", ""))
        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))
        if state.get("sig_time"):
            self.sig_time.setTime(QTime.fromString(state["sig_time"], "HH:mm"))
        # Restore card contents
        cards_state = state.get("cards", {})
        for key, content in cards_state.items():
            if key in self.cards:
                self.cards[key].set_content_text(content)

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[CTO4Form] Set patient name: {patient_info['name']}")
