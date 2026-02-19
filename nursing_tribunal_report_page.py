# ================================================================
#  NURSING TRIBUNAL REPORT PAGE — Nursing Tribunal Report Writer (T134)
# ================================================================

from __future__ import annotations

import re
import sys
from datetime import datetime, timedelta
from html import unescape

from PySide6.QtCore import Qt, Signal, QSize, QEvent
from PySide6.QtGui import QColor, QFontDatabase
from PySide6.QtWidgets import QGraphicsDropShadowEffect
from shared_widgets import create_zoom_row, add_lock_to_popup
from spell_check_textedit import enable_spell_check_on_textedit
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QSplitter, QStackedWidget, QTextEdit,
    QSizePolicy, QPushButton, QToolButton, QComboBox, QColorDialog, QMessageBox, QSlider,
    QApplication
)
from mypsy_richtext_editor import MyPsychAdminRichTextEditor
from utils.resource_path import resource_path


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# NURSING TOOLBAR (Same as Tribunal Toolbar)
# ================================================================

class NursingToolbar(QWidget):
    """Toolbar for the Nursing Tribunal Report Page."""

    # Formatting signals
    set_font_family = Signal(str)
    set_font_size = Signal(int)

    toggle_bold = Signal()
    toggle_italic = Signal()
    toggle_underline = Signal()

    set_text_color = Signal(QColor)
    set_highlight_color = Signal(QColor)

    set_align_left = Signal()
    set_align_center = Signal()
    set_align_right = Signal()
    set_align_justify = Signal()

    bullet_list = Signal()
    numbered_list = Signal()

    indent = Signal()
    outdent = Signal()

    undo = Signal()
    redo = Signal()

    insert_date = Signal()
    insert_section_break = Signal()

    export_docx = Signal()
    check_spelling = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(80)
        self.setStyleSheet("""
            NursingToolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 500;
            }
            QToolButton:hover {
                background: rgba(0,0,0,0.08);
            }
            QComboBox {
                background: rgba(255,255,255,0.85);
                color: #333333;
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 17px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #e0e0e0;
            }
        """)

        # Outer layout for the toolbar
        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        # Scroll area to prevent compression
        scroll = QScrollArea()
        scroll.setWidgetResizable(False)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        # Container widget for toolbar content
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        container.setMinimumWidth(1200)  # Force scrollbar when viewport is smaller
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 8, 2)
        layout.setSpacing(10)

        # ---------------------------------------------------------
        # EXPORT BUTTON - PROMINENT STYLING
        # ---------------------------------------------------------
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(160, 38)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 17px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #1d4ed8;
            }
            QToolButton:pressed {
                background: #1e40af;
            }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # ---------------------------------------------------------
        # UPLOADED DOCS BUTTON (dropdown menu)
        # ---------------------------------------------------------
        from PySide6.QtWidgets import QMenu
        import_btn = QToolButton()
        import_btn.setText("Uploaded Docs")
        import_btn.setFixedSize(160, 38)
        import_btn.setPopupMode(QToolButton.InstantPopup)
        self.upload_menu = QMenu()
        import_btn.setMenu(self.upload_menu)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 17px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #059669;
            }
            QToolButton:pressed {
                background: #047857;
            }
            QToolButton::menu-indicator { image: none; }
        """)
        layout.addWidget(import_btn)

        # ---------------------------------------------------------
        # FONT FAMILY
        # ---------------------------------------------------------
        self.font_combo = QComboBox()
        self.font_combo.setFixedWidth(180)

        families = QFontDatabase.families()
        if sys.platform == "win32":
            preferred = ["Segoe UI", "Calibri", "Cambria", "Arial", "Times New Roman"]
        else:
            preferred = ["Avenir Next", "Avenir", "SF Pro Text", "Helvetica Neue", "Helvetica"]

        added = set()
        for f in preferred:
            if f in families:
                self.font_combo.addItem(f)
                added.add(f)
        for f in families:
            if f not in added:
                self.font_combo.addItem(f)

        self.font_combo.currentTextChanged.connect(self.set_font_family.emit)
        layout.addWidget(self.font_combo)

        # ---------------------------------------------------------
        # FONT SIZE
        # ---------------------------------------------------------
        self.size_combo = QComboBox()
        self.size_combo.setFixedWidth(65)
        for sz in [8, 9, 10, 11, 12, 14, 16, 18, 20, 22]:
            self.size_combo.addItem(str(sz))
        self.size_combo.currentTextChanged.connect(lambda v: self.set_font_size.emit(int(v)))
        layout.addWidget(self.size_combo)

        # Simple button helper - fixed size to prevent compression
        def btn(label, slot):
            b = QToolButton()
            b.setText(label)
            b.setMinimumWidth(36)
            b.clicked.connect(slot)
            return b

        # ---------------------------------------------------------
        # BASIC STYLES
        # ---------------------------------------------------------
        layout.addWidget(btn("B", self.toggle_bold.emit))
        layout.addWidget(btn("I", self.toggle_italic.emit))
        layout.addWidget(btn("U", self.toggle_underline.emit))

        # ---------------------------------------------------------
        # COLORS
        # ---------------------------------------------------------
        layout.addWidget(btn("A", self._choose_text_color))
        layout.addWidget(btn("🖍", self._choose_highlight_color))

        # ---------------------------------------------------------
        # ALIGNMENT
        # ---------------------------------------------------------
        layout.addWidget(btn("L", self.set_align_left.emit))
        layout.addWidget(btn("C", self.set_align_center.emit))
        layout.addWidget(btn("R", self.set_align_right.emit))
        layout.addWidget(btn("J", self.set_align_justify.emit))

        # ---------------------------------------------------------
        # LISTS / INDENTATION
        # ---------------------------------------------------------
        layout.addWidget(btn("•", self.bullet_list.emit))
        layout.addWidget(btn("1.", self.numbered_list.emit))
        layout.addWidget(btn("→", self.indent.emit))
        layout.addWidget(btn("←", self.outdent.emit))

        # ---------------------------------------------------------
        # UNDO / REDO
        # ---------------------------------------------------------
        layout.addWidget(btn("⟲", self.undo.emit))
        layout.addWidget(btn("⟳", self.redo.emit))

        # ---------------------------------------------------------
        # INSERTS
        # ---------------------------------------------------------
        layout.addWidget(btn("Date", self.insert_date.emit))
        layout.addWidget(btn("Break", self.insert_section_break.emit))

        # ---------------------------------------------------------
        # SPELL CHECK
        # ---------------------------------------------------------
        spell_btn = QToolButton()
        spell_btn.setText("Spell Check")
        spell_btn.setFixedSize(120, 38)
        spell_btn.setStyleSheet("""
            QToolButton {
                background: #f59e0b;
                color: white;
                font-size: 16px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 6px 12px;
            }
            QToolButton:hover { background: #d97706; }
            QToolButton:pressed { background: #b45309; }
        """)
        spell_btn.setToolTip("Jump to next spelling error")
        spell_btn.clicked.connect(self.check_spelling.emit)
        layout.addWidget(spell_btn)

        # Finalize scroll area
        scroll.setWidget(container)
        outer_layout.addWidget(scroll)

    # -----------------------------
    # COLOR PICKERS
    # -----------------------------
    def _choose_text_color(self):
        col = QColorDialog.getColor(QColor("black"), self)
        if col.isValid():
            self.set_text_color.emit(col)

    def _choose_highlight_color(self):
        col = QColorDialog.getColor(QColor("yellow"), self)
        if col.isValid():
            self.set_highlight_color.emit(col)


# ================================================================
# NURSING CARD WIDGET
# ================================================================

class NursingCardWidget(QFrame):
    """A clickable card for a nursing tribunal report section."""

    clicked = Signal(str)  # Emits the section key

    STYLE_NORMAL = """
        NursingCardWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
            padding: 16px;
        }
        NursingCardWidget:hover {
            border-color: #14b8a6;
            background: #f0fdfa;
        }
    """

    STYLE_SELECTED = """
        NursingCardWidget {
            background: #ccfbf1;
            border: 2px solid #14b8a6;
            border-left: 4px solid #0d9488;
            border-radius: 12px;
            padding: 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)

        # Allow card to expand horizontally to fill container
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Title
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 19px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        layout.addWidget(title_lbl)

        # Editor (rich text with formatting support)
        self.editor = MyPsychAdminRichTextEditor()
        self.editor.setPlaceholderText("Click to edit...")
        self.editor.setReadOnly(False)
        self._editor_height = 100
        self.editor.setMinimumHeight(60)
        self.editor.setMaximumHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
                color: #374151;
            }
        """)
        editor_zoom = create_zoom_row(self.editor, base_size=12)
        layout.addLayout(editor_zoom)
        layout.addWidget(self.editor)

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(12)
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #14b8a6;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        """Handle drag events on the expand bar."""
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(60, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setMinimumHeight(self._editor_height)
                self.editor.setMaximumHeight(self._editor_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        # Only emit if not clicking on the editor or expand bar
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        """Set the selected state of the card."""
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        """Return whether the card is selected."""
        return self._selected

    def mousePressEvent(self, event):
        # Only emit if not clicking on the editor or expand bar
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)


# ================================================================
# NURSING HEADING WIDGET (for Yes/No sections with no card content)
# ================================================================

class NursingHeadingWidget(QFrame):
    """A clickable heading for nursing sections that don't need an editor card."""

    clicked = Signal(str)  # Emits the section key

    STYLE_NORMAL = """
        NursingHeadingWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 8px 16px;
        }
        NursingHeadingWidget:hover {
            border-color: #14b8a6;
            background: #f0fdfa;
        }
    """

    STYLE_SELECTED = """
        NursingHeadingWidget {
            background: #ccfbf1;
            border: 2px solid #14b8a6;
            border-left: 4px solid #0d9488;
            border-radius: 8px;
            padding: 8px 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.setMinimumHeight(50)

        # Add subtle shadow
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(12)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 25))
        self.setGraphicsEffect(shadow)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)

        # Title - with word wrap enabled
        self.title_lbl = QLabel(title)
        self.title_lbl.setWordWrap(True)
        self.title_lbl.setStyleSheet("""
            font-size: 17px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        self.title_lbl.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        layout.addWidget(self.title_lbl, 1)

        # Click indicator arrow
        arrow = QLabel("\u25B6")
        arrow.setStyleSheet("""
            font-size: 12px;
            color: #14b8a6;
            background: transparent;
        """)
        arrow.setFixedWidth(20)
        arrow.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        layout.addWidget(arrow)

        # Dummy editor for compatibility
        self.editor = _NursingDummyEditor()

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        return self._selected


class _NursingDummyEditor:
    """Minimal dummy editor for NursingHeadingWidget compatibility."""
    def toPlainText(self):
        return ""
    def setPlainText(self, text):
        pass
    def toHtml(self):
        return ""
    def setHtml(self, html):
        pass
    def clear(self):
        pass


# ================================================================
# FIXED DATA PANEL FOR NURSING (Incidents, etc.)
# ================================================================

class NursingFixedDataPanel(QWidget):
    """A panel for displaying fixed data from extraction."""

    sent = Signal(str)  # Emits content to send to card

    def __init__(self, title: str, subtitle: str = "", parent=None):
        super().__init__(parent)
        self._title = title
        self._entries = []
        self.notes = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(12)

        # Title
        title_label = QLabel(title)
        title_label.setStyleSheet("""
            font-size: 20px;
            font-weight: 600;
            color: #0f766e;
        """)
        layout.addWidget(title_label)

        if subtitle:
            subtitle_label = QLabel(subtitle)
            subtitle_label.setStyleSheet("""
                font-size: 17px;
                color: #6b7280;
            """)
            layout.addWidget(subtitle_label)

        # Date info
        self.date_info = QLabel("")
        self.date_info.setStyleSheet("""
            font-size: 16px;
            color: #9ca3af;
            font-style: italic;
        """)
        layout.addWidget(self.date_info)

        # Scroll area for entries
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setStyleSheet("""
            QScrollArea {
                background: transparent;
                border: none;
            }
        """)

        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(12, 12, 12, 12)
        self.content_layout.setSpacing(8)
        scroll.setWidget(self.content_widget)

        layout.addWidget(scroll, 1)

        # Summary frame (hidden by default)
        self.summary_frame = QFrame()
        self.summary_frame.setStyleSheet("""
            QFrame {
                background: #fef3c7;
                border: 1px solid #f59e0b;
                border-radius: 8px;
                padding: 8px;
            }
        """)
        self.summary_frame.hide()
        summary_layout = QVBoxLayout(self.summary_frame)
        summary_layout.setContentsMargins(8, 8, 8, 8)

        summary_header = QHBoxLayout()
        summary_label = QLabel("Summary")
        summary_label.setStyleSheet("font-weight: 600; color: #92400e;")
        summary_header.addWidget(summary_label)
        summary_header.addStretch()

        copy_summary_btn = QPushButton("Copy")
        copy_summary_btn.setFixedWidth(50)
        copy_summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 4px 8px;
                border-radius: 4px;
                font-size: 16px;
            }
            QPushButton:hover {
                background: #d97706;
            }
        """)
        copy_summary_btn.clicked.connect(self._copy_summary)
        summary_header.addWidget(copy_summary_btn)

        close_btn = QPushButton("X")
        close_btn.setFixedWidth(24)
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #92400e;
                border: none;
                font-size: 19px;
                font-weight: bold;
            }
            QPushButton:hover {
                color: #78350f;
                background: rgba(0,0,0,0.1);
                border-radius: 4px;
            }
        """)
        close_btn.clicked.connect(lambda: self.summary_frame.hide())
        summary_header.addWidget(close_btn)

        summary_layout.addLayout(summary_header)

        self.summary_text = QTextEdit()
        self.summary_text.setMaximumHeight(150)
        self.summary_text.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #fcd34d;
                border-radius: 4px;
                font-size: 17px;
            }
        """)
        summary_layout.addWidget(self.summary_text)
        layout.addWidget(self.summary_frame)

        # Button row
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #f3f4f6;
                color: #374151;
                border: 1px solid #d1d5db;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #e5e7eb;
            }
        """)
        clear_btn.clicked.connect(self._clear)
        btn_layout.addWidget(clear_btn)

        self.summary_btn = QPushButton("Summary")
        self.summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #d97706;
            }
        """)
        self.summary_btn.clicked.connect(self._generate_summary)
        btn_layout.addWidget(self.summary_btn)

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #14b8a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #0d9488;
            }
        """)
        send_btn.clicked.connect(self._send_to_letter)
        btn_layout.addWidget(send_btn)

        layout.addLayout(btn_layout)

    def set_entries(self, entries: list, date_range_info: str = ""):
        """Set the entries to display."""
        self._entries = entries

        if date_range_info:
            self.date_info.setText(date_range_info)

        # Clear existing content
        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # For incident panels, show summary directly instead of all entries
        if "harm" in self._title.lower() or "property" in self._title.lower() or "seclusion" in self._title.lower():
            self._show_incident_summary_directly()
            return

        for entry in entries[:100]:  # Limit display to match psychiatric report
            entry_frame = QFrame()
            entry_frame.setStyleSheet("""
                QFrame {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 6px;
                    padding: 8px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(8, 6, 8, 6)
            entry_layout.setSpacing(4)

            date_val = entry.get('date') or entry.get('datetime') or 'Unknown date'
            # Convert datetime to string if needed
            if hasattr(date_val, 'strftime'):
                date_str = date_val.strftime('%d/%m/%Y')
            else:
                date_str = str(date_val)
            date_label = QLabel(date_str)
            date_label.setStyleSheet("font-size: 16px; color: #6b7280; font-weight: 500;")
            entry_layout.addWidget(date_label)

            content = entry.get('content', '') or entry.get('text', '')
            preview = content[:200] + "..." if len(content) > 200 else content
            content_label = QLabel(preview)
            content_label.setWordWrap(True)
            content_label.setStyleSheet("font-size: 17px; color: #374151;")
            entry_layout.addWidget(content_label)

            self.content_layout.addWidget(entry_frame)

        self.content_layout.addStretch()

    def _show_incident_summary_directly(self):
        """Show filtered incident summary directly in popup."""
        # Group entries by date
        date_to_entries = {}
        for entry in self._entries:
            date_str = entry.get('date') or entry.get('datetime') or 'Unknown date'
            content = entry.get('content', '') or entry.get('text', '')
            if not content.strip():
                continue

            # Normalize date
            try:
                if isinstance(date_str, str):
                    for fmt in ['%d/%m/%Y', '%Y-%m-%d', '%d %b %Y', '%d-%m-%Y']:
                        try:
                            parsed_date = datetime.strptime(date_str.strip()[:10], fmt)
                            date_str = parsed_date.strftime('%d/%m/%Y')
                            break
                        except:
                            continue
            except:
                pass

            if date_str not in date_to_entries:
                date_to_entries[date_str] = []
            date_to_entries[date_str].append(content.strip())

        # Filter patterns (same as psychiatric report)
        def should_filter_line(line: str) -> bool:
            line_lower = line.lower().strip()

            # Filter lines starting with certain prefixes
            filter_prefixes = ['diagnosis:', 'positive behaviour', 'to self:', 'to others:',
                             'risk:', 'risks', 'self neglect:']
            for prefix in filter_prefixes:
                if line_lower.startswith(prefix):
                    return True

            # Filter lines containing certain phrases
            filter_phrases = ['without incident', 'no evidence', 'nothing to indicate',
                            'risk of', 'risk to', 'medication for agitation',
                            'call police if', 'less agitation', 'less aggression',
                            'reduced agitation', 'reduced aggression',
                            'police and ambulance to be called if', 'police to be called if',
                            '(agitation)', 'previous', 'threatened to walk out',
                            'can be aggressive']
            for phrase in filter_phrases:
                if phrase in line_lower:
                    return True

            # Word boundary filtering for nil, non, no
            import re
            if re.search(r'\bnil\b', line_lower):
                return True
            if re.search(r'\bnon\b', line_lower):
                return True
            if re.search(r'\bno\b', line_lower):
                return True

            return False

        # Build filtered output
        seen_content = set()
        output_lines = []

        sorted_dates = sorted(date_to_entries.keys(), key=lambda d: datetime.strptime(d, '%d/%m/%Y') if '/' in d else datetime.min, reverse=True)

        for date_str in sorted_dates:
            entries = date_to_entries[date_str]
            filtered_entries = []

            for content in entries:
                lines = content.split('\n')
                filtered_lines = []
                for line in lines:
                    if line.strip() and not should_filter_line(line):
                        normalized = ' '.join(line.lower().split())
                        if normalized not in seen_content:
                            seen_content.add(normalized)
                            filtered_lines.append(line.strip())
                if filtered_lines:
                    filtered_entries.append(' '.join(filtered_lines))

            if filtered_entries:
                combined = ' | '.join(filtered_entries)
                output_lines.append(f"{date_str}: {combined}")

        # Display in popup
        count = len(output_lines)
        self.date_info.setText(f"{count} incident(s) after filtering")
        self.summary_btn.hide()  # Hide summary button for incident panels

        summary_frame = QFrame()
        summary_frame.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px;
            }
        """)
        summary_layout = QVBoxLayout(summary_frame)

        summary_text = QTextEdit()
        summary_text.setPlainText('\n'.join(output_lines) if output_lines else "No incidents found after filtering.")
        summary_text.setStyleSheet("""
            QTextEdit {
                background: #fafafa;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                font-size: 17px;
                padding: 8px;
            }
        """)
        summary_text.setMinimumHeight(200)
        summary_layout.addWidget(summary_text)

        self.content_layout.addWidget(summary_frame)
        self.content_layout.addStretch()

        # Store for send to letter
        self.notes = output_lines

    def _clear(self):
        """Clear all entries."""
        self._entries = []
        self.notes = []
        while self.content_layout.count():
            child = self.content_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        self.summary_frame.hide()

    def _copy_summary(self):
        """Copy summary to clipboard."""
        from PySide6.QtWidgets import QApplication
        clipboard = QApplication.clipboard()
        clipboard.setText(self.summary_text.toPlainText())

    def _generate_summary(self):
        """Generate smart summary based on panel type."""
        if not self._entries:
            self.summary_text.setPlainText("No data to summarize.")
            self.summary_frame.show()
            return

        # For Progress panel, use risk-based summary with engagement section
        if "Progress" in self._title or "Engagement" in self._title:
            summary_parts = self._generate_risk_based_summary()
            if summary_parts:
                self.summary_text.setPlainText('\n'.join(summary_parts))
                self.summary_frame.show()
                return

        # Default: simple summary
        summary_lines = []
        for entry in self._entries[:20]:
            date_val = entry.get('date') or entry.get('datetime') or ''
            if hasattr(date_val, 'strftime'):
                date_str = date_val.strftime('%d/%m/%Y')
            else:
                date_str = str(date_val)
            content = entry.get('content', '') or entry.get('text', '')
            preview = content[:100] + "..." if len(content) > 100 else content
            summary_lines.append(f"{date_str}: {preview}")

        self.summary_text.setPlainText('\n'.join(summary_lines))
        self.summary_frame.show()

    def _generate_risk_based_summary(self):
        """Generate narrative summary with Progress, Engagement, and Insight sections."""
        import re
        from datetime import datetime
        from pathlib import Path

        if not self._entries:
            return []

        # Load risk dictionary for scoring
        risk_dict = {}
        risk_file = Path(__file__).parent / "riskDICT.txt"
        if risk_file.exists():
            with open(risk_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or ',' not in line:
                        continue
                    parts = line.rsplit(',', 1)
                    term = parts[0].strip().lower()
                    score_str = parts[1].strip() if len(parts) > 1 else ''
                    if term and score_str:
                        try:
                            risk_dict[term] = int(score_str)
                        except:
                            pass

        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        def get_highest_scoring_line(content, risk_dict):
            """Return the line with the highest risk score (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            best_line = None
            best_score = -1
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                line_score = sum(points for term, points in risk_dict.items() if term in line_lower)
                if line_score > best_score:
                    best_score = line_score
                    best_line = cleaned
            return best_line

        def get_relevant_keyword_line(content, keywords):
            """Extract the most relevant line containing keywords (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                # Skip headings
                if line_lower.endswith(':') and len(line_lower) < 40:
                    continue
                for kw in keywords:
                    if kw in line_lower:
                        # Remove common heading prefixes
                        for prefix in ['Observations and Engagement:', 'Engagement:', 'Self-care:', 'Self care:', 'Insight:']:
                            if cleaned.startswith(prefix):
                                cleaned = cleaned[len(prefix):].strip()
                        if len(cleaned) >= 100:
                            return cleaned
            return None

        def extract_specific_line(content, keywords):
            """Extract the specific line containing any of the keywords (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                # Skip headings like "Observations and Engagement:"
                if line_lower.endswith(':') and len(line_lower) < 40:
                    continue
                for kw in keywords:
                    if kw in line_lower:
                        # Remove common heading prefixes
                        for prefix in ['Observations and Engagement:', 'Engagement:', 'Self-care:', 'Self care:', 'Insight:']:
                            if cleaned.startswith(prefix):
                                cleaned = cleaned[len(prefix):].strip()
                        if len(cleaned) >= 100:
                            return cleaned
            return None

        # Categorize entries
        entries_data = []
        for entry in self._entries:
            content = entry.get('content', '') or entry.get('text', '') or ''
            content_lower = content.lower()
            if not content:
                continue

            note_date = parse_date(entry.get('date') or entry.get('datetime'))
            date_str = note_date.strftime('%d/%m/%Y') if note_date else 'Unknown'

            # Calculate score
            score = sum(points for term, points in risk_dict.items() if term in content_lower)

            entries_data.append({
                'date': note_date,
                'date_str': date_str,
                'score': score,
                'content': content,
                'content_lower': content_lower
            })

        # Sort by date
        entries_data.sort(key=lambda x: x['date'] or datetime.min)

        # === EXTRACT EVENTS BY CATEGORY ===

        # High concern events (threshold 1500)
        concern_events = [e for e in entries_data if e['score'] >= 1500]

        # Violence/Aggression events
        violence_keywords = ['violence', 'violent', 'assault', 'attack', 'fight', 'aggression', 'aggressive', 'physical altercation']
        violence_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in violence_keywords)]

        # DNA / Lack of contact
        dna_keywords = ['dna', 'did not attend', 'no contact', 'uncontactable', 'failed to attend', 'missed appointment']
        dna_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in dna_keywords)]

        # Positive progress
        positive_keywords = ['stable', 'settled', 'calm', 'pleasant', 'appropriate', 'well presented', 'good rapport', 'engaging well', 'cooperative', 'progress']
        positive_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in positive_keywords) and e['score'] < 100]

        # Self-care keywords - specific hygiene/appearance terms only (exclude sleep)
        selfcare_keywords = ['well kempt', 'good self care', 'good self-care', 'poor self care', 'poor self-care',
                            'well dressed', 'showered', 'washed', 'malodorous', 'smelt', 'clean', 'unkempt',
                            'dishevelled', 'neglected appearance', 'poor hygiene']
        selfcare_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in selfcare_keywords)]

        # Insight records
        insight_keywords = ['insight', 'awareness', 'understands', 'accepts diagnosis', 'denies illness', 'lacks insight', 'good insight', 'poor insight', 'partial insight']
        insight_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in insight_keywords)]

        # === ENGAGEMENT / ACTIVITIES SECTION ===
        # General engagement keywords
        engagement_keywords = ['engage', 'engaged', 'engagement', 'did not engage', 'no engagement', 'refused to engage', 'engaged well', 'good engagement', 'poor engagement', 'minimal engagement']
        engagement_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in engagement_keywords)]

        # OT, Groups, Section 17 leave
        ot_keywords = ['occupational therapy', 'ot session', 'ot group', 'o.t.', 'occupational therapist']
        ot_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in ot_keywords)]

        groups_keywords = ['group', 'groups', 'attended group', 'psychology group', 'therapy group', 'art group', 'music group', 'community meeting', 'ward round']
        groups_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in groups_keywords)]

        s17_keywords = ['section 17', 's17', 'leave', 'escorted leave', 'unescorted leave', 'ground leave', 'community leave', 'home leave', 'overnight leave']
        s17_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in s17_keywords)]

        # === BUILD NARRATIVE SUMMARY ===
        summary_parts = []

        # Date range header
        dates = [e['date'] for e in entries_data if e['date']]
        if dates:
            earliest = min(dates).strftime('%d/%m/%Y')
            latest = max(dates).strftime('%d/%m/%Y')

        # === SECTION 1: PROGRESS ===
        summary_parts.append("PROGRESS")
        summary_parts.append("-" * 40)

        if dates:
            summary_parts.append(f"Review period: {earliest} to {latest}")

        # Concern events - show all with highest scoring line
        if concern_events:
            summary_parts.append("")
            summary_parts.append(f"Significant behavioural concerns noted ({len(concern_events)}):")
            for e in sorted(concern_events, key=lambda x: x['date'] or datetime.min, reverse=True):
                relevant_line = get_highest_scoring_line(e['content'], risk_dict)
                if relevant_line:
                    summary_parts.append(f"  * {e['date_str']}: {relevant_line}")
        else:
            summary_parts.append("")
            summary_parts.append("No significant behavioural concerns in this period.")

        # Violence/Aggression - show all high score with relevant line
        summary_parts.append("")
        summary_parts.append("Episodes of violence or aggression:")
        high_violence = [e for e in violence_events if e['score'] >= 2000]
        if high_violence:
            summary_parts.append(f"  ({len(high_violence)} episodes)")
            for e in sorted(high_violence, key=lambda x: x['date'] or datetime.min, reverse=True):
                relevant_line = get_highest_scoring_line(e['content'], risk_dict)
                if relevant_line:
                    summary_parts.append(f"  * {e['date_str']}: {relevant_line}")
        else:
            summary_parts.append("  Nil noted.")

        # Positive progress - summarize to key descriptors only
        def summarize_positive(content_lower):
            """Summarize positive entry to key descriptor."""
            if 'superficially settled' in content_lower or 'superfically settled' in content_lower:
                return 'settled (superficially)'
            if 'settled' in content_lower or 'calm' in content_lower:
                return 'settled'
            if 'stable' in content_lower:
                return 'stable'
            if 'pleasant' in content_lower:
                return 'pleasant'
            if 'cooperative' in content_lower:
                return 'cooperative'
            if 'engaging well' in content_lower or 'good rapport' in content_lower:
                return 'engaging well'
            if 'well presented' in content_lower or 'appropriate' in content_lower:
                return 'well presented'
            if 'progress' in content_lower:
                return 'progress noted'
            return 'positive'

        if positive_events:
            sorted_positive = sorted(positive_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append("")
            summary_parts.append(f"Positive observations ({len(sorted_positive)}):")
            for e in sorted_positive:
                descriptor = summarize_positive(e['content_lower'])
                summary_parts.append(f"  * {e['date_str']}: {descriptor}")

        # === SECTION 2: SELF-CARE ===
        summary_parts.append("")
        summary_parts.append("")
        summary_parts.append("SELF-CARE")
        summary_parts.append("-" * 40)
        if selfcare_events:
            sorted_selfcare = sorted(selfcare_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_selfcare)} records)")
            for e in sorted_selfcare:
                specific_line = extract_specific_line(e['content'], selfcare_keywords)
                if specific_line:
                    summary_parts.append(f"  * {e['date_str']}: {specific_line}")
        else:
            summary_parts.append("  No specific self-care records found.")

        # === SECTION 3: ENGAGEMENT / ACTIVITIES ===
        summary_parts.append("")
        summary_parts.append("")
        summary_parts.append("ENGAGEMENT / ACTIVITIES")
        summary_parts.append("-" * 40)

        # Summarization function for engagement
        def summarize_engagement(content_lower):
            """Summarize engagement entry to key descriptor."""
            if 'did not engage' in content_lower or 'no engagement' in content_lower or 'nil engagement' in content_lower:
                return 'no engagement'
            if 'refused to engage' in content_lower or 'declined to engage' in content_lower:
                return 'refused to engage'
            if 'minimal engagement' in content_lower or 'limited engagement' in content_lower:
                return 'minimal engagement'
            if 'engaged well' in content_lower or 'good engagement' in content_lower:
                return 'engaged well'
            if 'engaged on needs' in content_lower or 'needs basis' in content_lower or 'needs led' in content_lower or 'needy bas' in content_lower:
                return 'engaged on needs basis'
            if 'polite when engag' in content_lower or 'pleasant when engag' in content_lower:
                return 'polite when engaged'
            if 'lack of engagement' in content_lower or 'unable to assess' in content_lower:
                return 'lack of engagement'
            if 'engaged' in content_lower:
                return 'engaged'
            return 'engagement noted'

        # General Engagement
        summary_parts.append("")
        summary_parts.append("General Engagement:")
        if engagement_events:
            sorted_engagement = sorted(engagement_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_engagement)} records)")
            for e in sorted_engagement:
                descriptor = summarize_engagement(e['content_lower'])
                summary_parts.append(f"  * {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No engagement records found.")

        # Summarization function for OT
        def summarize_ot(content_lower):
            """Summarize OT entry to key descriptor."""
            if 'refused' in content_lower or 'declined' in content_lower:
                return 'declined OT'
            if 'engaged well' in content_lower or 'participated' in content_lower:
                return 'engaged in OT'
            if 'ot session' in content_lower or 'occupational therapy' in content_lower:
                return 'OT session'
            return 'OT noted'

        # OT
        summary_parts.append("")
        summary_parts.append("Occupational Therapy:")
        if ot_events:
            sorted_ot = sorted(ot_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_ot)} records)")
            for e in sorted_ot:
                descriptor = summarize_ot(e['content_lower'])
                summary_parts.append(f"  * {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No OT sessions recorded.")

        # Summarization function for groups
        def summarize_groups(content_lower):
            """Summarize groups entry to key descriptor."""
            if 'did not attend' in content_lower or 'refused' in content_lower or 'declined' in content_lower:
                return 'did not attend'
            if 'ward round' in content_lower:
                return 'ward round'
            if 'community meeting' in content_lower:
                return 'community meeting'
            if 'psychology group' in content_lower or 'therapy group' in content_lower:
                return 'therapy group'
            if 'art group' in content_lower:
                return 'art group'
            if 'music group' in content_lower:
                return 'music group'
            if 'attended group' in content_lower or 'group session' in content_lower:
                return 'attended group'
            if 'group room' in content_lower:
                return 'group room session'
            return 'group activity'

        # Groups
        summary_parts.append("")
        summary_parts.append("Groups / Ward Activities:")
        if groups_events:
            sorted_groups = sorted(groups_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_groups)} records)")
            for e in sorted_groups:
                descriptor = summarize_groups(e['content_lower'])
                summary_parts.append(f"  * {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No group activities recorded.")

        # Summarization function for S17 leave
        def summarize_leave(content_lower):
            """Summarize S17 leave entry to key descriptor."""
            # Check for NON-USE of leave FIRST (most specific patterns)
            if ('leave: none' in content_lower or 'leave: nil' in content_lower or
                'leave: not utilised' in content_lower or 'leave: not used' in content_lower or
                'leave: non used' in content_lower or 'leave not utilised' in content_lower or
                'leave not used' in content_lower or 'did not use leave' in content_lower or
                'no leave today' in content_lower or 'leave not taken' in content_lower or
                'did not utilise leave' in content_lower or 'none utilised' in content_lower or
                'did not access leave' in content_lower or 'no leave' in content_lower):
                return 'leave not used'
            if 'refused leave' in content_lower or 'declined leave' in content_lower:
                return 'refused leave'
            if 'leave could not' in content_lower or 'unable to facilitate' in content_lower:
                return 'leave not facilitated'
            # Check for ACTUAL USE - must have clear indicators of leave being taken
            # Shopping leave - but check it was actually used
            if ('shopping leave' in content_lower or 'went shopping' in content_lower or
                'to the shop' in content_lower or 'marston green' in content_lower or
                'chelmsley wood' in content_lower) and 'leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'shopping leave'
            # Community leave - check it was used
            if 'community leave' in content_lower:
                if 'went on community' in content_lower or 'utilised community' in content_lower or 'accessed community' in content_lower:
                    return 'community leave'
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'community leave'
            # Ground leave - check it was used
            if 'ground leave' in content_lower or 'grounds leave' in content_lower:
                if ('went on' in content_lower or 'utilised' in content_lower or 'accessed' in content_lower or
                    'facilitated' in content_lower or 'took' in content_lower or 'had' in content_lower):
                    return 'ground leave'
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'ground leave'
            # Escorted leave - check it was used
            if 'escorted leave' in content_lower:
                if ('went on' in content_lower or 'utilised' in content_lower or 'accessed' in content_lower or
                    'facilitated' in content_lower or 'took' in content_lower):
                    return 'escorted leave'
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'escorted leave'
            # Unescorted leave
            if 'unescorted leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'unescorted leave'
            # Home leave
            if 'home leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'home leave'
            # Overnight leave
            if 'overnight leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'overnight leave'
            # Clear indicators that leave was actually taken
            if ('went on leave' in content_lower or 'utilised leave' in content_lower or
                'accessed leave' in content_lower or 'leave was facilitated' in content_lower or
                'leave went well' in content_lower or 'took him on leave' in content_lower or
                'took her on leave' in content_lower or 'escorted on leave' in content_lower or
                'used his leave' in content_lower or 'used her leave' in content_lower or
                'utilised his leave' in content_lower or 'utilised her leave' in content_lower or
                'leave to' in content_lower and ('shop' in content_lower or 'walk' in content_lower)):
                return 'leave utilised'
            # Default: if just "leave" mentioned without clear use indicator, assume not used
            return 'leave not used'

        # Section 17 Leave
        summary_parts.append("")
        summary_parts.append("Section 17 Leave:")
        if s17_events:
            sorted_s17 = sorted(s17_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_s17)} records)")
            for e in sorted_s17:
                descriptor = summarize_leave(e['content_lower'])
                summary_parts.append(f"  * {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No Section 17 leave recorded.")

        # === SECTION 4: INSIGHT ===
        summary_parts.append("")
        summary_parts.append("")
        summary_parts.append("INSIGHT")
        summary_parts.append("-" * 40)
        if insight_events:
            sorted_insight = sorted(insight_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_insight)} records)")
            for e in sorted_insight:
                specific_line = extract_specific_line(e['content'], insight_keywords)
                if specific_line:
                    summary_parts.append(f"  * {e['date_str']}: {specific_line}")
        else:
            summary_parts.append("  No specific insight assessments found.")

        return summary_parts

    def _send_to_letter(self):
        """Send content to the card."""
        if self.notes:
            self.sent.emit('\n'.join(self.notes))
        elif self._entries:
            content_parts = []
            for entry in self._entries[:10]:
                date_str = entry.get('date') or entry.get('datetime') or ''
                content = entry.get('content', '') or entry.get('text', '')
                if content:
                    content_parts.append(f"{date_str}: {content[:200]}")
            self.sent.emit('\n'.join(content_parts))


# ================================================================
# MAIN NURSING TRIBUNAL REPORT PAGE
# ================================================================

class NursingTribunalReportPage(QWidget):
    """Main page for creating Nursing Tribunal Reports (T134)."""

    go_back = Signal()  # Signal to return to reports page

    # Sections based on T134 form - Nursing Report
    SECTIONS = [
        ("1. Patient Details", "patient_details"),
        ("2. Factors affecting understanding or ability to cope with hearing", "factors_hearing"),
        ("3. Adjustments for tribunal to consider", "adjustments"),
        ("4. Nature of nursing care and medication", "nursing_care"),
        ("5. Level of observation", "observation_level"),
        ("6. Contact with relatives, friends or other patients", "contact"),
        ("7. Community support", "community_support"),
        ("8. Strengths or positive factors", "strengths"),
        ("9. Current progress, engagement, behaviour, cooperation, activities, self-care and insight", "progress"),
        ("10. AWOL or failed return from leave", "awol"),
        ("11. Compliance with medication/treatment", "compliance"),
        ("12. Incidents of harm to self or others", "risk_harm"),
        ("13. Incidents of property damage", "risk_property"),
        ("14. Seclusion or restraint", "seclusion"),
        ("15. Section 2: Detention justified for health, safety or protection", "s2_detention"),
        ("16. Other sections: Medical treatment justified for health, safety or protection", "other_detention"),
        ("17. Risk if discharged from hospital", "discharge_risk"),
        ("18. Community risk management", "community"),
        ("19. Other relevant information", "other_info"),
        ("20. Recommendations to tribunal", "recommendations"),
        ("21. Signature", "signature"),
    ]

    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self.cards = {}
        self.popups = {}
        self.popup_memory = {}
        self._active_editor = None
        self._current_gender = "Male"
        self._selected_card_key = None
        self._my_details = self._load_my_details()

        # Store extracted data
        self._extracted_raw_notes = []
        self._extracted_categories = {}
        self._incident_data = []

        # Guard flags to prevent reprocessing on navigation
        self._data_processed_id = None
        self._notes_processed_id = None

        self._setup_ui()

        # Connect to shared store for cross-talk with psychiatric form
        self._connect_shared_store()

    def _load_my_details(self) -> dict:
        """Load clinician details from database."""
        if not self.db:
            return {}

        details = self.db.get_clinician_details()
        if not details:
            return {}

        return {
            "full_name": details[1] or "",
            "role_title": details[2] or "",
            "discipline": details[3] or "",
            "registration_body": details[4] or "",
            "registration_number": details[5] or "",
            "phone": details[6] or "",
            "email": details[7] or "",
            "team_service": details[8] or "",
            "hospital_org": details[9] or "",
            "ward_department": details[10] or "",
            "signature_block": details[11] or "",
        }

    def _connect_shared_store(self):
        """Connect to SharedDataStore for cross-report data sharing."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.report_sections_changed.connect(self._on_report_sections_changed)
            shared_store.notes_changed.connect(self._on_notes_changed)
            shared_store.extracted_data_changed.connect(self._on_extracted_data_changed)
            shared_store.patient_info_changed.connect(self._on_patient_info_changed)
            print("[NURSING] Connected to SharedDataStore signals (sections, notes, extracted_data, patient_info)")

            # Check if there's already data in the store
            self._check_shared_store_for_existing_data()
        except Exception as e:
            print(f"[NURSING] Failed to connect to SharedDataStore: {e}")

    def _check_shared_store_for_existing_data(self):
        """Check SharedDataStore for existing data when page is created."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()

            # Check for existing report sections (cross-talk)
            existing_sections = shared_store.report_sections
            source = shared_store.get_report_source()
            if existing_sections and source != "nursing_tribunal":
                print(f"[NURSING] Found existing sections from {source}, populating...")
                self._on_report_sections_changed(existing_sections, source)

            # Check for existing notes (only if no report data)
            if not self._has_report_data():
                notes = shared_store.notes
                if notes:
                    print(f"[NURSING] Found {len(notes)} existing notes in SharedDataStore")
                    if not hasattr(self, '_extracted_raw_notes'):
                        self._extracted_raw_notes = []
                    self._extracted_raw_notes = notes

                # Check for existing extracted data
                extracted_data = shared_store.extracted_data
                if extracted_data:
                    print(f"[NURSING] Found existing extracted data in SharedDataStore")
                    self._on_extracted_data_changed(extracted_data)
        except Exception as e:
            print(f"[NURSING] Error checking shared store: {e}")

    def _on_patient_info_changed(self, patient_info: dict):
        """Handle patient info updates from SharedDataStore."""
        if patient_info and any(patient_info.values()):
            print(f"[NURSING] Received patient info from SharedDataStore: {list(k for k,v in patient_info.items() if v)}")
            self._fill_patient_details(patient_info)

    def _has_report_data(self):
        """Check if report data has been imported (local or via SharedDataStore)."""
        if hasattr(self, '_imported_report_data') and self._imported_report_data:
            return True
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            source = shared_store.get_report_source()
            if source and source != "nursing_tribunal" and shared_store.report_sections:
                return True
        except Exception:
            pass
        return False

    def _on_notes_changed(self, notes: list):
        """Handle notes updates from SharedDataStore."""
        if notes:
            # Skip if report data exists (report takes priority)
            if self._has_report_data():
                print(f"[NURSING] Skipping notes from SharedDataStore - report data already imported")
                return
            print(f"[NURSING] Received {len(notes)} notes from SharedDataStore")
            if not hasattr(self, '_extracted_raw_notes'):
                self._extracted_raw_notes = []
            self._extracted_raw_notes = notes

    def _on_extracted_data_changed(self, data: dict):
        """Handle extracted data updates from SharedDataStore."""
        if not data:
            return
        # Skip if report data exists (report takes priority)
        if self._has_report_data():
            print(f"[NURSING] Skipping extracted data from SharedDataStore - report data already imported")
            return
        print(f"[NURSING] Received extracted data from SharedDataStore: {list(data.keys())}")
        if not hasattr(self, '_extracted_categories'):
            self._extracted_categories = {}
        categories = data.get("categories", data)
        self._extracted_categories = categories

    def _on_report_sections_changed(self, sections: dict, source_form: str):
        """Handle report sections imported from another form (cross-talk)."""
        # Only process if from psychiatric form (not our own import)
        if source_form == "nursing_tribunal":
            return

        print(f"[NURSING] Cross-talk received from {source_form}: {len(sections)} sections")

        # Store imported data for popups to use
        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        # Populate matching sections
        loaded_count = 0
        for key, content in sections.items():
            if key in self.cards and content:
                # Store the imported data for the popup's imported data section
                self._imported_report_data[key] = content
                loaded_count += 1
                print(f"[NURSING] Cross-talk stored: {key}")

                # If popup already exists, populate it with report data
                if hasattr(self, 'popups') and key in self.popups:
                    popup = self.popups[key]
                    self._populate_single_popup(popup, key, content)
                    print(f"[NURSING] Populated popup '{key}' from cross-talk")

        if loaded_count > 0:
            print(f"[NURSING] Cross-talk populated {loaded_count} sections from {source_form}")

    def _check_shared_store_for_sections(self):
        """Check SharedDataStore for existing sections when form is shown."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            existing_sections = shared_store.report_sections
            source = shared_store.get_report_source()
            if existing_sections and source and source != "nursing_tribunal":
                print(f"[NURSING] showEvent: Found sections from {source}")
                self._on_report_sections_changed(existing_sections, source)
        except Exception as e:
            print(f"[NURSING] Error checking shared store: {e}")

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar (teal color for nursing)
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("""
            QFrame {
                background: #0f766e;
                border-bottom: 1px solid #0d9488;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        # Back button
        back_btn = QPushButton("< Back")
        back_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: white;
                border: 1px solid rgba(255,255,255,0.3);
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: rgba(255,255,255,0.1);
            }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Nursing Tribunal Report")
        title.setStyleSheet("""
            font-size: 24px;
            font-weight: 700;
            color: white;
        """)
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Clear Report button
        clear_btn = QPushButton("Clear Report - Start New")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover {
                background: #7f1d1d;
            }
            QPushButton:pressed {
                background: #450a0a;
            }
        """)
        clear_btn.clicked.connect(self._clear_report)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = NursingToolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        # Connect uploaded docs menu to SharedDataStore
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())
        # Helper to get active editor (persists when toolbar clicked)
        def get_active_editor():
            return self._active_editor

        # Helper to safely call editor method
        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        # Connect formatting signals
        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(
                        self,
                        "Spell Check",
                        "No spelling errors found."
                    )

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area with splitter
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("""
            QScrollArea {
                background: #f3f4f6;
                border: none;
            }
        """)
        self.main_splitter.addWidget(self.cards_holder)

        self.editor_root = QWidget()
        self.editor_root.setStyleSheet("background: #f3f4f6;")
        self.editor_root.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.editor_layout = QVBoxLayout(self.editor_root)
        self.editor_layout.setContentsMargins(24, 24, 24, 24)
        self.editor_layout.setSpacing(16)
        self.cards_holder.setWidget(self.editor_root)

        # Right: Panel (resizable)
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(350)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame {
                background: rgba(245,245,245,0.98);
                border-left: 1px solid rgba(0,0,0,0.08);
            }
        """)
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 450])

        # Connect splitter movement to update card widths
        self.main_splitter.splitterMoved.connect(self._on_splitter_moved)

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        # Header row with title and lock button
        self.panel_header = QWidget()
        self.panel_header.setStyleSheet("""
            background: #0f766e;
            border-radius: 8px;
            margin-bottom: 8px;
        """)
        header_layout = QHBoxLayout(self.panel_header)
        header_layout.setContentsMargins(16, 12, 16, 12)
        header_layout.setSpacing(8)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 24px;
            font-weight: 700;
            color: white;
            background: transparent;
        """)
        header_layout.addWidget(self.panel_title, 1)

        # Lock button in header
        self.header_lock_btn = QPushButton("Unlocked")
        self.header_lock_btn.setFixedSize(70, 26)
        self.header_lock_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.header_lock_btn.setToolTip("Click to lock this section")
        self.header_lock_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255, 255, 255, 0.3);
                border: 2px solid rgba(255, 255, 255, 0.8);
                border-radius: 13px;
                font-size: 13px;
                font-weight: 600;
                color: white;
            }
            QPushButton:hover { background: rgba(255, 255, 255, 0.5); }
        """)
        self.header_lock_btn.clicked.connect(self._toggle_current_popup_lock)
        self.header_lock_btn.hide()
        header_layout.addWidget(self.header_lock_btn)

        panel_layout.addWidget(self.panel_header)

        # Popup stack
        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        # Create all cards
        self._create_cards()

    # Sections that should be headings (no editor card, just clickable title)
    HEADING_ONLY_SECTIONS = {
        "s2_detention",      # Section 15
        "other_detention",   # Section 16
    }

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _toggle_current_popup_lock(self):
        """Toggle lock on the currently active popup."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'toggle_lock'):
            popup.toggle_lock()
            self._update_header_lock_button()

    def _update_header_lock_button(self):
        """Update header lock button to match current popup state."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'is_locked'):
            is_locked = popup.is_locked()
            if is_locked:
                self.header_lock_btn.setText("Locked")
                self.header_lock_btn.setToolTip("Click to unlock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(239, 68, 68, 0.5);
                        border: 2px solid #ef4444;
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: white;
                    }
                    QPushButton:hover { background: rgba(239, 68, 68, 0.7); }
                """)
            else:
                self.header_lock_btn.setText("Unlocked")
                self.header_lock_btn.setToolTip("Click to lock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(255, 255, 255, 0.3);
                        border: 2px solid rgba(255, 255, 255, 0.8);
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: white;
                    }
                    QPushButton:hover { background: rgba(255, 255, 255, 0.5); }
                """)
            self.header_lock_btn.show()
        else:
            self.header_lock_btn.hide()

    def _set_current_popup(self, popup):
        """Set the current active popup and update lock button."""
        self._current_popup = popup
        self._update_header_lock_button()

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self  # Capture reference to self for closure

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _create_cards(self):
        """Create all section cards (or headings for certain sections)."""
        for title, key in self.SECTIONS:
            if key in self.HEADING_ONLY_SECTIONS:
                # Create heading widget instead of full card
                card = NursingHeadingWidget(title, key, parent=self.editor_root)
            else:
                # Create standard card widget
                card = NursingCardWidget(title, key, parent=self.editor_root)
                # Hook up focus event to register this editor as active
                self._hook_editor_focus(card.editor)
                # Connect card text changes to sync with popup
                card.editor.textChanged.connect(lambda k=key: self._on_card_text_changed(k))
            card.clicked.connect(self._on_card_clicked)
            self.cards[key] = card
            self.editor_layout.addWidget(card)

        self.editor_layout.addStretch()

        # Prefill signature card from my_details on startup
        self._prefill_signature_card()

    def _prefill_signature_card(self):
        """Prefill signature card from clinician details."""
        if not self._my_details:
            return

        from tribunal_popups import SignaturePopup
        popup = SignaturePopup(parent=self, my_details=self._my_details)
        popup.hide()  # Hide immediately - we just need the text
        if hasattr(popup, 'generate_text'):
            text = popup.generate_text()
            if text and text.strip():
                self.cards["signature"].editor.setPlainText(text)
                print(f"[NURSING] Prefilled signature card on startup")
        popup.deleteLater()  # Clean up the popup

    def _on_splitter_moved(self, pos, index):
        """Update card widths when splitter is moved."""
        self._update_card_widths()

    def _update_card_widths(self):
        """Update card widths based on current viewport size."""
        viewport_width = self.cards_holder.viewport().width()
        # Account for margins (24px on each side)
        card_width = viewport_width - 48
        if card_width > 100:  # Minimum sensible width
            for key, card in self.cards.items():
                if key in self.HEADING_ONLY_SECTIONS:
                    # Headings use max width so text wraps
                    card.setMaximumWidth(card_width)
                else:
                    card.setFixedWidth(card_width)

    def showEvent(self, event):
        """Set initial card widths when page is shown."""
        super().showEvent(event)
        # Use a timer to ensure layout is complete
        from PySide6.QtCore import QTimer
        QTimer.singleShot(50, self._update_card_widths)

        # Check for existing report sections in shared store (cross-talk)
        QTimer.singleShot(100, self._check_shared_store_for_sections)

    def resizeEvent(self, event):
        """Update card widths on window resize."""
        super().resizeEvent(event)
        self._update_card_widths()

    def _on_card_clicked(self, key: str):
        """Handle card click - show appropriate popup."""
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        print(f"[NURSING] Card clicked: {key}")

        # Update card selection
        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)
        if key in self.cards:
            self.cards[key].setSelected(True)
        self._selected_card_key = key

        # Create popup if not exists
        if key not in self.popups:
            popup = self._create_popup(key)
            if popup:
                self.popups[key] = popup
                self.popup_stack.addWidget(popup)
                popup.sent.connect(lambda text, k=key: self._update_card(k, text))

                # Connect gender_changed signal for patient_details popup
                if key == "patient_details" and hasattr(popup, 'gender_changed'):
                    popup.gender_changed.connect(self._on_gender_changed)

                # Populate from imported report data if available
                if hasattr(self, '_imported_report_data') and key in self._imported_report_data:
                    if not getattr(popup, '_imported_data_added', False):
                        content = self._imported_report_data[key]
                        self._populate_single_popup(popup, key, content)
                        print(f"[NURSING] Populated popup '{key}' from imported report data")

                # Notes-based searches - only run when importing notes, NOT report imports
                _has_report = hasattr(self, '_imported_report_data') and self._imported_report_data
                if not _has_report:
                    # Populate Seclusion panel with all raw notes for seclusion/restraint search (section 14)
                    if key == "seclusion" and self._extracted_raw_notes and hasattr(popup, 'set_notes'):
                        popup.set_notes(self._extracted_raw_notes)
                        print(f"[NURSING] Seclusion popup searching {len(self._extracted_raw_notes)} notes")

                    # Populate Risk Harm panel with all raw notes for harm search (section 12)
                    if key == "risk_harm" and self._extracted_raw_notes and hasattr(popup, 'set_notes'):
                        popup.set_notes(self._extracted_raw_notes)
                        print(f"[NURSING] Risk Harm popup searching {len(self._extracted_raw_notes)} notes")

                    # Populate Risk Property panel with all raw notes for property damage search (section 13)
                    if key == "risk_property" and self._extracted_raw_notes and hasattr(popup, 'set_notes'):
                        popup.set_notes(self._extracted_raw_notes)
                        print(f"[NURSING] Risk Property popup searching {len(self._extracted_raw_notes)} notes")

                    # Populate AWOL panel with all raw notes for AWOL search
                    if key == "awol" and self._extracted_raw_notes and hasattr(popup, 'set_notes'):
                        popup.set_notes(self._extracted_raw_notes)
                        print(f"[NURSING] AWOL popup searching {len(self._extracted_raw_notes)} notes")

                    # Populate Compliance panel with all raw notes for non-compliance search
                    if key == "compliance" and self._extracted_raw_notes and hasattr(popup, 'set_notes'):
                        popup.set_notes(self._extracted_raw_notes)
                        print(f"[NURSING] Compliance popup searching {len(self._extracted_raw_notes)} notes")

                    # Populate Discharge Risk panel (section 17) with notes for risk analysis - same as psych tribunal section 21
                    if key == "discharge_risk" and self._extracted_raw_notes and hasattr(popup, 'set_notes_for_risk_analysis'):
                        popup.set_notes_for_risk_analysis(self._extracted_raw_notes)
                        print(f"[NURSING] Discharge Risk popup analyzing {len(self._extracted_raw_notes)} notes for risk")

                # Prefill signature card if popup has content from my_details
                if key == "signature" and hasattr(popup, 'generate_text'):
                    text = popup.generate_text()
                    if text and text.strip():
                        self._update_card(key, text)
                        print(f"[NURSING] Prefilled signature card from my_details")

                # Populate Progress panel (section 9) with notes - same as psych tribunal section 14
                if not _has_report and key == "progress" and self._extracted_raw_notes and hasattr(popup, 'set_entries'):
                    # Use recent notes (last 6 months) if available from pending data, else use raw notes
                    progress_entries = self._pending_section_data.get("progress", self._extracted_raw_notes[:100]) if hasattr(self, '_pending_section_data') else self._extracted_raw_notes[:100]
                    if progress_entries:
                        popup.set_entries(progress_entries, f"{len(progress_entries)} notes")
                        print(f"[NURSING] Progress popup populated with {len(progress_entries)} notes")

                # Apply any pending data for this section (notes pipeline only)
                if not _has_report and hasattr(self, '_pending_section_data') and key in self._pending_section_data:
                    pending_entries = self._pending_section_data[key]
                    if pending_entries:
                        # risk_harm, risk_property, and seclusion use set_notes (TribunalRisk*Popup/TribunalSeclusionPopup)
                        if key in ["risk_harm", "risk_property", "seclusion"] and hasattr(popup, 'set_notes'):
                            popup.set_notes(pending_entries)
                            print(f"[NURSING] Applied pending data to '{key}' ({len(pending_entries)} notes for search)")
                        # discharge_risk uses set_notes_for_risk_analysis (GPRRiskPopup)
                        elif key == "discharge_risk" and hasattr(popup, 'set_notes_for_risk_analysis'):
                            popup.set_notes_for_risk_analysis(pending_entries)
                            print(f"[NURSING] Applied pending data to '{key}' ({len(pending_entries)} notes for risk analysis)")
                        elif hasattr(popup, 'set_entries'):
                            # Use specific info text for progress section
                            if key == "progress":
                                info_text = f"{len(pending_entries)} most recent notes"
                            else:
                                info_text = f"{len(pending_entries)} entries"
                            popup.set_entries(pending_entries, info_text)
                            print(f"[NURSING] Applied pending data to '{key}' ({len(pending_entries)} entries)")

                # Pre-fill medications for nursing_care section (section 4)
                if key == "nursing_care" and self._extracted_raw_notes:
                    self._prefill_medications_from_notes()

        if key in self.popups:
            popup = self.popups[key]
            self.popup_stack.setCurrentWidget(popup)
            self._set_current_popup(popup)

            # Send popup form content to card (imported data checkbox is unchecked so only form data flows)
            if hasattr(self, '_imported_report_data') and key in self._imported_report_data:
                if hasattr(popup, '_send_to_card'):
                    popup._send_to_card()

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details popup and card from extracted demographics."""
        if not patient_info:
            return

        # Create popup if it doesn't exist
        if "patient_details" not in self.popups:
            from tribunal_popups import PatientDetailsPopup
            popup = PatientDetailsPopup(parent=self)
            self.popups["patient_details"] = popup
            self.popup_stack.addWidget(popup)

        # Fill the popup fields
        popup = self.popups["patient_details"]
        if hasattr(popup, 'fill_patient_info'):
            popup.fill_patient_info(patient_info)

            # Also update the card preview
            text = popup.generate_text()
            if "patient_details" in self.cards and text.strip():
                self.cards["patient_details"].editor.setPlainText(text)
                print(f"[NursingTribunalReport] Updated patient_details card with demographics")

    def _create_popup(self, key: str):
        """Create the appropriate popup for a section."""
        from tribunal_popups import PatientDetailsPopup, FactorsHearingPopup, AdjustmentsPopup
        from tribunal_popups import StrengthsPopup, CompliancePopup, YesNoNAPopup
        from tribunal_popups import DischargeRiskPopup, CommunityManagementPopup
        from tribunal_popups import RecommendationsPopup, SignaturePopup, SimpleYesNoPopup

        if key == "patient_details":
            return PatientDetailsPopup(parent=self)
        elif key == "factors_hearing":
            return FactorsHearingPopup(parent=self, gender=self._current_gender)
        elif key == "adjustments":
            return AdjustmentsPopup(parent=self, gender=self._current_gender)
        elif key == "nursing_care":
            return NursingCarePopup(parent=self, gender=self._current_gender)
        elif key == "observation_level":
            return ObservationLevelPopup(parent=self)
        elif key == "contact":
            return ContactPopup(parent=self, gender=self._current_gender)
        elif key == "community_support":
            return CommunitySupportPopup(parent=self, gender=self._current_gender)
        elif key == "strengths":
            return StrengthsPopup(parent=self, gender=self._current_gender)
        elif key == "progress":
            # Section 9: Use TribunalProgressPopup identical to psych tribunal section 14
            # Filter to 1 year from most recent entry
            from tribunal_popups import TribunalProgressPopup
            return TribunalProgressPopup(parent=self, date_filter='1_year')
        elif key == "awol":
            # Section 10: Use CarePathwayPopup style identical to ASR section 11
            from social_tribunal_report_page import CarePathwayPopup
            return CarePathwayPopup(parent=self, gender=self._current_gender)
        elif key == "compliance":
            # Section 11: Use TribunalCompliancePopup with existing inputs + imported non-compliance data
            from tribunal_popups import TribunalCompliancePopup
            return TribunalCompliancePopup(parent=self, gender=self._current_gender)
        elif key == "risk_harm":
            # Section 12: Use TribunalRiskHarmPopup identical to psych tribunal section 17
            from tribunal_popups import TribunalRiskHarmPopup
            return TribunalRiskHarmPopup(parent=self)
        elif key == "risk_property":
            # Section 13: Use TribunalRiskPropertyPopup identical to psych tribunal section 18
            from tribunal_popups import TribunalRiskPropertyPopup
            return TribunalRiskPropertyPopup(parent=self)
        elif key == "seclusion":
            # Section 14: Use TribunalSeclusionPopup with imported seclusion/restraint data
            from tribunal_popups import TribunalSeclusionPopup
            return TribunalSeclusionPopup(parent=self)
        elif key == "s2_detention":
            # Section 15: Use SimpleYesNoPopup identical to psychiatric section 19
            from tribunal_popups import SimpleYesNoPopup
            return SimpleYesNoPopup(
                title="Section 2 Detention",
                question="Is detention under Section 2 justified for health, safety or protection of others?",
                parent=self
            )
        elif key == "other_detention":
            return SimpleYesNoPopup(
                title="Medical Treatment Justified",
                question="Is the provision of medical treatment in hospital justified or necessary in the interests of the patient's health or safety, or for the protection of others?",
                parent=self
            )
        elif key == "discharge_risk":
            # Section 17: Use GPRRiskPopup identical to psychiatric tribunal section 21
            from general_psychiatric_report_page import GPRRiskPopup
            return GPRRiskPopup(parent=self)
        elif key == "community":
            return CommunityManagementPopup(parent=self, gender=self._current_gender)
        elif key == "other_info":
            return OtherInfoPopup(parent=self)
        elif key == "recommendations":
            try:
                from general_psychiatric_report_page import GPRLegalCriteriaPopup
                from icd10_dict import ICD10_DICT
                print(f"[NURSING] Creating GPRLegalCriteriaPopup: ICD10_DICT={len(ICD10_DICT)} entries, gender={self._current_gender}")
                popup = GPRLegalCriteriaPopup(parent=self, gender=self._current_gender, icd10_dict=ICD10_DICT)
                print(f"[NURSING] GPRLegalCriteriaPopup created successfully")
                return popup
            except Exception as e:
                print(f"[NURSING] ERROR creating recommendations popup: {e}")
                import traceback
                traceback.print_exc()
                from PySide6.QtWidgets import QLabel
                err = QLabel(f"ERROR loading recommendations popup:\n{e}")
                err.setStyleSheet("color: red; font-size: 16px; padding: 20px;")
                err.setWordWrap(True)
                return err
        elif key == "signature":
            return SignaturePopup(parent=self, my_details=self._my_details)

        return None

    def _on_gender_changed(self, gender: str):
        """Handle gender change from patient details - update all popups."""
        self._current_gender = gender
        print(f"[NURSING] Gender changed to: {gender}")

        # Update all existing popups that support gender
        for key, popup in self.popups.items():
            if hasattr(popup, 'set_gender'):
                popup.set_gender(gender)
                print(f"[NURSING] Updated gender for popup: {key}")

    def _update_card(self, key: str, text: str):
        """Update a card with text from popup - REPLACES existing content."""
        if key in self.cards:
            self.cards[key].editor.setPlainText(text.strip())

    def _on_card_text_changed(self, key: str):
        """Handle card text changes - sync to popup if not already syncing."""
        if getattr(self, '_syncing', False):
            return

        if key not in self.cards:
            return

        card_text = self.cards[key].editor.toPlainText()

        # Parse card text and update corresponding popup
        if key in self.popups:
            self._syncing = True
            try:
                self._update_popup_from_card(key, card_text)
            finally:
                self._syncing = False

    def _update_popup_from_card(self, key: str, card_text: str):
        """Update popup fields from card text."""
        if key not in self.popups:
            return

        popup = self.popups[key]
        fields = self._parse_structured_text(card_text)

        # Patient details popup
        if key == "patient_details":
            # Handle name field
            name_value = fields.get('Full Name') or fields.get('Name') or ''
            if name_value and hasattr(popup, 'name_field'):
                popup.name_field.blockSignals(True)
                popup.name_field.setText(name_value)
                popup.name_field.blockSignals(False)

            # Handle DOB field
            dob_value = fields.get('Date of Birth') or fields.get('DOB') or ''
            if dob_value and hasattr(popup, 'dob_field'):
                popup.dob_field.blockSignals(True)
                from PySide6.QtCore import QDate
                for fmt in ["dd/MM/yyyy", "dd-MM-yyyy", "yyyy-MM-dd", "d MMMM yyyy"]:
                    parsed = QDate.fromString(dob_value.strip(), fmt)
                    if parsed.isValid():
                        popup.dob_field.setDate(parsed)
                        break
                popup.dob_field.blockSignals(False)

            # Handle Gender field
            gender_value = fields.get('Gender') or ''
            if gender_value and hasattr(popup, 'gender_field'):
                popup.gender_field.blockSignals(True)
                idx = popup.gender_field.findText(gender_value, Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    popup.gender_field.setCurrentIndex(idx)
                popup.gender_field.blockSignals(False)
                # Propagate gender since signals were blocked
                self._on_gender_changed(gender_value)

            # Handle Residence field
            residence_value = fields.get('Usual Place of Residence') or fields.get('Residence') or fields.get('Address') or ''
            if residence_value and hasattr(popup, 'residence_field'):
                popup.residence_field.blockSignals(True)
                popup.residence_field.setPlainText(residence_value)
                popup.residence_field.blockSignals(False)

    def _parse_structured_text(self, text: str) -> dict:
        """Parse text with 'Field: Value' format into a dictionary."""
        fields = {}
        for line in text.split('\n'):
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    field_name = parts[0].strip()
                    value = parts[1].strip()
                    if field_name and value:
                        fields[field_name] = value
        return fields

    def _go_back(self):
        """Return to reports page."""
        self.go_back.emit()

    def _clear_report(self):
        """Clear all data and start new report."""
        reply = QMessageBox.question(
            self,
            "Clear Report",
            "Are you sure you want to clear all data and start a new report?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            # Clear all cards
            for card in self.cards.values():
                if hasattr(card, 'editor'):
                    card.editor.clear()

            # Clear stored data
            self._extracted_raw_notes = []
            self._extracted_categories = {}
            self._incident_data = []
            self._data_processed_id = None
            self._notes_processed_id = None
            if hasattr(self, '_imported_report_data'):
                self._imported_report_data = {}
            if hasattr(self, '_imported_report_sections'):
                self._imported_report_sections = {}
            if hasattr(self, '_pending_section_data'):
                self._pending_section_data = {}

            # Destroy all popups and remove from stack
            for key, popup in list(self.popups.items()):
                if hasattr(self, 'popup_stack'):
                    self.popup_stack.removeWidget(popup)
                popup.deleteLater()
            self.popups.clear()

            # Clear data extractor
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                if hasattr(self._data_extractor_overlay, 'clear_extraction'):
                    self._data_extractor_overlay.clear_extraction()

            # Recreate signature popup so mydetails get restored
            if hasattr(self, '_on_card_clicked'):
                if "signature" in self.cards:
                    self._on_card_clicked("signature")

            print("[NURSING] Report cleared - ready for new report")

    def _export_docx(self):
        """Export report to DOCX using T134 template with table-based input boxes."""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from docx import Document
        from docx.shared import Pt
        from datetime import datetime
        import os
        import shutil

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Nursing Tribunal Report",
            f"nursing_tribunal_report_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            # Find template path - use new template with real input boxes
            template_path = resource_path("templates", "t134_template_new.docx")

            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Template Error", f"T134 template not found at:\n{template_path}")
                return

            # Copy template to destination
            shutil.copy(template_path, file_path)

            # Open the copied template
            doc = Document(file_path)

            # Helper to get card content (or popup state for heading-only sections)
            def get_content(key):
                # For heading-only sections, get content from popup's generate_text()
                if key in self.HEADING_ONLY_SECTIONS:
                    if key in self.popups and hasattr(self.popups[key], 'generate_text'):
                        return self.popups[key].generate_text()
                    return ""
                # For regular cards, get from editor
                if key in self.cards:
                    return self.cards[key].editor.toPlainText().strip()
                return ""

            # Helper to extract detail text after Yes/No prefix
            def extract_detail(content):
                if not content:
                    return ""
                detail = content
                lower = detail.lower()
                for prefix in ["yes, ", "yes - ", "yes-", "yes,", "yes.", "no, ", "no - ", "no-", "no,", "no."]:
                    if lower.startswith(prefix):
                        detail = detail[len(prefix):].strip()
                        break
                if lower == "yes" or lower == "no":
                    return ""
                return detail

            # Helper to set table cell text
            def set_cell(table_idx, row, col, text):
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    if row < len(table.rows) and col < len(table.rows[row].cells):
                        cell = table.cell(row, col)
                        cell.text = text
                        # Set font size for all paragraphs in cell
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(11)

            # ============================================================
            # PARSE SIGNATURE INFO
            # ============================================================
            sig_content = get_content("signature")
            sig_lines = sig_content.split('\n') if sig_content else []
            sig_name = sig_role = ""
            sig_date = datetime.now().strftime("%d/%m/%Y")

            for line in sig_lines:
                line_lower = line.lower()
                if "signed:" in line_lower or "name:" in line_lower:
                    sig_name = line.split(":", 1)[-1].strip()
                elif "designation:" in line_lower or "role:" in line_lower:
                    sig_role = line.split(":", 1)[-1].strip()
                elif "date:" in line_lower:
                    sig_date = line.split(":", 1)[-1].strip()

            # ============================================================
            # TABLE 0: Header - nested table (7x1) in right cell
            # Structure: Row 0=Hospital logo title, Row 1=Your name title,
            #            Row 2=name entry, Row 3=Your role title,
            #            Row 4=role entry, Row 5=Date title, Row 6=date entry
            # ============================================================
            header_table = doc.tables[0]
            right_cell = header_table.cell(0, 1)
            if right_cell.tables:
                info_table = right_cell.tables[0]
                # Row 2: Your name entry
                if len(info_table.rows) > 2 and sig_name:
                    info_table.rows[2].cells[0].text = sig_name
                # Row 4: Your role entry
                if len(info_table.rows) > 4 and sig_role:
                    info_table.rows[4].cells[0].text = sig_role
                # Row 6: Date entry
                if len(info_table.rows) > 6 and sig_date:
                    info_table.rows[6].cells[0].text = sig_date

            # ============================================================
            # PATIENT DETAILS
            # ============================================================
            pd_content = get_content("patient_details")
            pd_lines = pd_content.split('\n') if pd_content else []
            full_name = dob = residence = ""
            for line in pd_lines:
                line_lower = line.lower()
                if "name:" in line_lower or "full name:" in line_lower:
                    full_name = line.split(":", 1)[-1].strip()
                elif "date of birth:" in line_lower or "dob:" in line_lower:
                    dob = line.split(":", 1)[-1].strip()
                elif "residence:" in line_lower or "place of residence:" in line_lower or "usual place" in line_lower:
                    residence = line.split(":", 1)[-1].strip()

            # TABLE 1: Full name (1x2)
            set_cell(1, 0, 1, full_name)

            # TABLE 2: DOB character boxes (1x8)
            if dob and len(doc.tables) > 2:
                dob_table = doc.tables[2]
                dob_clean = dob.replace("/", "").replace("-", "").replace(" ", "")
                for i, char in enumerate(dob_clean[:len(dob_table.columns)]):
                    dob_table.cell(0, i).text = char

            # TABLE 3: Residence (1x1)
            set_cell(3, 0, 0, residence)

            # ============================================================
            # HELPER: Check if content indicates Yes answer
            # ============================================================
            def has_yes_content(content):
                """Returns True if content indicates Yes answer."""
                if not content:
                    return False
                lower = content.lower().strip()
                # Check for various "No" patterns
                no_patterns = [
                    "no ", "no,", "no-", "no.",  # Starts with no
                    "there have been no",        # Seclusion: "There have been no occasions..."
                    "there is no other",         # Other info: "There is no other relevant..."
                    "no occasions",              # Alternative no occasion pattern
                    "does not have",             # Contact: "does not have contact"
                    "no factors",                # Section 2: "no factors"
                    "no adjustments",            # Section 3: "no adjustments"
                    "not applicable",            # N/A pattern
                    "n/a",                       # N/A short
                ]
                if lower == "no":
                    return False
                for pattern in no_patterns:
                    if pattern in lower:
                        return False
                # Any other content means Yes
                return True

            # ============================================================
            # SECTION INPUT BOXES (Tables 4-20)
            # Sections 15 & 16 are checkbox-only (no entry table)
            # ============================================================
            table_mapping = {
                4: ("factors_hearing", True),      # Section 2
                5: ("adjustments", True),          # Section 3
                6: ("nursing_care", False),        # Section 4
                7: ("observation_level", False),   # Section 5
                8: ("contact", True),              # Section 6
                9: ("community_support", False),   # Section 7
                10: ("strengths", False),          # Section 8
                11: ("progress", False),           # Section 9
                12: ("awol", False),               # Section 10
                13: ("compliance", False),         # Section 11
                14: ("risk_harm", False),          # Section 12
                15: ("risk_property", False),      # Section 13
                16: ("seclusion", True),           # Section 14
                # Sections 15 & 16 are checkbox-only - handled below
                17: ("discharge_risk", True),      # Section 17
                18: ("community", False),          # Section 18
                19: ("other_info", True),          # Section 19
                20: ("recommendations", True),     # Section 20
            }

            # Sections with Yes/No checkboxes - only fill table if Yes
            yes_no_sections = {"seclusion", "factors_hearing", "adjustments", "contact",
                               "discharge_risk", "other_info", "recommendations"}

            for table_idx, (card_key, use_extract) in table_mapping.items():
                content = get_content(card_key)
                if content:
                    # For Yes/No sections, only fill table if answer is Yes
                    if card_key in yes_no_sections:
                        if not has_yes_content(content):
                            continue  # Skip table entry for No answers
                    if use_extract:
                        content = extract_detail(content) or content
                    set_cell(table_idx, 0, 0, content)

            # ============================================================
            # CHECKBOX SECTIONS - mark ☐ → ☒ based on content
            # ============================================================
            # Map section number prefixes to card keys
            # "Are there" is the start of section 2's question
            checkbox_sections = {
                "Are there are any factors": "factors_hearing",  # Section 2
                "2.": "factors_hearing",                          # Section 2 alt
                "3.": "adjustments",          # Section 3
                "6.": "contact",              # Section 6
                "14.": "seclusion",           # Section 14
                "15.": "s2_detention",        # Section 15 (checkbox only)
                "16.": "other_detention",     # Section 16 (checkbox only)
                "17.": "discharge_risk",      # Section 17
                "19.": "other_info",          # Section 19
                "20.": "recommendations",     # Section 20
            }

            current_section_key = None
            for para in doc.paragraphs:
                text = para.text.strip()

                # Detect which section we're in
                for sec_prefix, card_key in checkbox_sections.items():
                    if text.startswith(sec_prefix):
                        current_section_key = card_key
                        break

                # Mark checkboxes based on content
                if current_section_key and "☐" in para.text:
                    content = get_content(current_section_key)
                    is_yes = has_yes_content(content)

                    # Check if this is a Yes or No checkbox line
                    lower_text = para.text.lower()
                    is_yes_line = "yes" in lower_text
                    is_no_line = lower_text.strip().startswith("☐ no") or (lower_text.strip().startswith("☐") and "no" in lower_text and "yes" not in lower_text)

                    if is_yes_line and is_yes:
                        # Mark Yes checkbox
                        for run in para.runs:
                            if "☐" in run.text:
                                run.text = run.text.replace("☐", "☒", 1)
                                break
                    elif is_no_line and not is_yes and content:
                        # Mark No checkbox (only if explicitly No)
                        for run in para.runs:
                            if "☐" in run.text:
                                run.text = run.text.replace("☐", "☒", 1)
                                break

            # ============================================================
            # TABLE 21: Signature box (1x1)
            # ============================================================
            from docx.shared import Inches

            sig_text_parts = []
            if sig_name:
                sig_text_parts.append(sig_name)
            if sig_role:
                sig_text_parts.append(sig_role)

            for line in sig_lines:
                line_lower = line.lower()
                if "qualification" in line_lower:
                    sig_text_parts.append(line.split(":", 1)[-1].strip())
                elif "registration:" in line_lower or "gmc:" in line_lower or "nmc:" in line_lower:
                    reg = line.split(":", 1)[-1].strip()
                    sig_text_parts.append(f"Registration: {reg}")

            # Put signature image and text in Table 21
            if 21 < len(doc.tables):
                sig_cell = doc.tables[21].cell(0, 0)
                sig_cell.text = ""  # Clear existing content

                # Add signature image if exists
                sig_image_path = os.path.expanduser("~/MyPsychAdmin/signature.png")
                if os.path.exists(sig_image_path):
                    sig_para = sig_cell.paragraphs[0]
                    sig_run = sig_para.add_run()
                    sig_run.add_picture(sig_image_path, width=Inches(1.5))
                    # Add new paragraph for text
                    sig_cell.add_paragraph("\n".join(sig_text_parts))
                else:
                    sig_cell.text = "\n".join(sig_text_parts)

            # ============================================================
            # TABLE 22: Date character boxes (1x8) - ddmmyyyy format
            # ============================================================
            if sig_date and len(doc.tables) > 22:
                date_table = doc.tables[22]

                # Try to parse date into ddmmyyyy format
                date_clean = ""
                import re

                # Try various date formats
                date_formats = [
                    ("%d/%m/%Y", None),      # 17/01/2026
                    ("%d-%m-%Y", None),      # 17-01-2026
                    ("%d.%m.%Y", None),      # 17.01.2026
                    ("%d %B %Y", None),      # 17 January 2026
                    ("%d %b %Y", None),      # 17 Jan 2026
                    ("%Y-%m-%d", None),      # 2026-01-17 (ISO)
                    ("%B %d, %Y", None),     # January 17, 2026
                    ("%b %d, %Y", None),     # Jan 17, 2026
                ]

                parsed_date = None
                for fmt, _ in date_formats:
                    try:
                        parsed_date = datetime.strptime(sig_date.strip(), fmt)
                        break
                    except:
                        continue

                if parsed_date:
                    date_clean = parsed_date.strftime("%d%m%Y")
                else:
                    # Fallback: strip non-digits
                    date_clean = re.sub(r'[^0-9]', '', sig_date)
                    if len(date_clean) == 6:
                        # ddmmyy -> ddmmyyyy
                        year_part = date_clean[4:6]
                        year = "20" + year_part if int(year_part) < 50 else "19" + year_part
                        date_clean = date_clean[:4] + year

                # Fill character boxes
                for i, char in enumerate(date_clean[:8]):
                    if i < len(date_table.columns):
                        date_table.cell(0, i).text = char

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Report exported to:\n{file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export report:\n{str(e)}")

    def _is_tribunal_report(self, file_path: str) -> tuple:
        """Detect if a document is a tribunal report vs clinical notes.

        Returns: (is_report: bool, report_type: str, full_text: str)

        Tribunal reports are characterized by:
        - Numbered sections (1., 2., 3., etc.)
        - Specific tribunal headings (factors affecting hearing, recommendations, etc.)
        - Formal structure with patient details section

        Clinical notes are characterized by:
        - Date-based entries
        - Progress notes, observations
        - Free-form narrative
        """
        import re

        try:
            if file_path.lower().endswith('.docx'):
                from docx import Document
                doc = Document(file_path)
                full_text = ' '.join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text += ' ' + cell.text
            elif file_path.lower().endswith('.pdf'):
                # For PDFs, try to extract text
                try:
                    import subprocess
                    result = subprocess.run(['pdftotext', '-layout', file_path, '-'],
                                          capture_output=True, text=True, timeout=30)
                    full_text = result.stdout if result.returncode == 0 else ""
                except:
                    full_text = ""
            else:
                with open(file_path, 'r', errors='ignore') as f:
                    full_text = f.read()
        except Exception as e:
            print(f"[NURSING] Error reading file for detection: {e}")
            return (False, 'unknown', "")

        lower = full_text.lower()

        # Tribunal report indicators
        report_indicators = [
            r'\b\d+\.\s+(?:are there|what is|does the|give details|has the)',  # Numbered questions
            r'tribunal\s+report',
            r'mental\s+health\s+tribunal',
            r'factors\s+(?:that\s+)?(?:may\s+)?affect.*(?:hearing|understanding)',
            r'recommendations?\s+to\s+(?:the\s+)?tribunal',
            r'(?:nursing|psychiatric|social\s+circumstances?)\s+report',
            r't131|t132|t133|t134',
            r'responsible\s+clinician',
            r'approved\s+mental\s+health\s+professional',
        ]

        report_score = sum(1 for pattern in report_indicators if re.search(pattern, lower))

        # Clinical notes indicators
        notes_indicators = [
            r'\d{1,2}[/.-]\d{1,2}[/.-]\d{2,4}\s*[-:]\s*\w',  # Date entries
            r'progress\s+note',
            r'ward\s+round',
            r'mdt\s+meeting',
            r'nursing\s+(?:entry|note|observation)',
            r'mental\s+state\s+examination',
            r'risk\s+assessment\s+completed',
            r'patient\s+(?:seen|reviewed|assessed)',
        ]

        notes_score = sum(1 for pattern in notes_indicators if re.search(pattern, lower))

        print(f"[NURSING] Detection scores - Report: {report_score}, Notes: {notes_score}")

        # Determine report type
        report_type = self._detect_report_type(full_text)

        # If report indicators dominate, it's a tribunal report
        if report_score >= 2 and report_score > notes_score:
            return (True, report_type, full_text)

        # If notes indicators dominate, it's clinical notes
        if notes_score > report_score:
            return (False, 'notes', full_text)

        # If we detected a specific report type, trust that
        if report_type in ['T131', 'T134', 'social']:
            return (True, report_type, full_text)

        # Default to notes if unclear
        return (False, 'unknown', full_text)

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            from shared_data_store import get_shared_store
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file - auto-detect tribunal reports vs clinical notes."""
        # First, detect if this is a tribunal report or clinical notes
        is_report, report_type, full_text = self._is_tribunal_report(file_path)

        print(f"[NURSING] Import detection: is_report={is_report}, type={report_type}")

        if is_report:
            # If it's a PTR (T131) or social report, send straight to data extractor
            # which uses PTR_TO_NTR_MAP for correct cross-report section mapping
            if report_type in ('T131', 'social'):
                print(f"[NURSING] Cross-report import ({report_type}) - sending to data extractor for PTR->NTR mapping")
                self._send_to_data_extractor(file_path)
                return

            # It's a nursing (T134) report - use report parsing
            if file_path.lower().endswith('.pdf'):
                try:
                    from pdf_loader import load_tribunal_pdf

                    result = load_tribunal_pdf(file_path)

                    if result.get('sections'):
                        self._populate_from_pdf(result, file_path)
                        return
                    else:
                        # Try OCR for scanned PDFs
                        print("[NURSING] PDF has no text - trying OCR...")
                        ocr_result = self._parse_scanned_pdf(file_path)
                        if ocr_result and ocr_result.get('sections'):
                            self._populate_from_pdf(ocr_result, file_path)
                            return

                except Exception as e:
                    print(f"[NURSING] PDF report load error: {e}")

            elif file_path.lower().endswith('.docx'):
                try:
                    result = self._parse_tribunal_docx(file_path)
                    if result and result.get('sections'):
                        self._populate_from_docx(result, file_path)
                        return
                except Exception as e:
                    print(f"[NURSING] DOCX report load error: {e}")

            # If report parsing failed, fall through to notes import
            print("[NURSING] Report parsing failed - falling back to notes import")

        # It's clinical notes or report parsing failed - send to data extractor
        print("[NURSING] Using data extractor for notes import")
        self._send_to_data_extractor(file_path)

    def _detect_report_type(self, full_text: str) -> str:
        """Detect the type of tribunal report from document text.

        Returns: 'T134' (nursing), 'T131' (psychiatric), 'social', or 'unknown'
        """
        lower = full_text.lower()

        # Check for specific report type indicators
        if 'nursing report' in lower or 'in-patient nursing' in lower or 't134' in lower:
            return 'T134'
        elif 'social circumstances' in lower or 'social supervisor' in lower:
            return 'social'
        elif 'psychiatric report' in lower or 'responsible clinician' in lower or 't131' in lower:
            return 'T131'

        # Check for profession-specific keywords
        nursing_keywords = ['nursing care', 'observation level', 'seclusion', 'restraint', 'nurse']
        medical_keywords = ['diagnosis', 'medication', 'treatment', 'prognosis', 'clinician', 'psychiatrist']
        social_keywords = ['housing', 'accommodation', 'care plan', 'care pathway', 'social worker', 'amhp']

        nursing_score = sum(1 for kw in nursing_keywords if kw in lower)
        medical_score = sum(1 for kw in medical_keywords if kw in lower)
        social_score = sum(1 for kw in social_keywords if kw in lower)

        max_score = max(nursing_score, medical_score, social_score)
        if max_score == 0:
            return 'unknown'

        if nursing_score == max_score:
            return 'T134'
        elif social_score == max_score:
            return 'social'
        else:
            return 'T131'

    def _get_heading_to_card_mapping(self) -> list:
        """Get heading patterns mapped to nursing card keys.

        Returns a list of (pattern, card_key) tuples for matching.
        Patterns are matched against heading text to determine the section.
        This enables cross-talk between T131, T133, T134 reports.

        IMPORTANT: Match by heading CONTENT, never by section numbers.
        """
        return [
            # === COMMON SECTIONS (appear in all report types) ===

            # Factors affecting hearing/understanding
            # T131/T133/T134: "Are there any factors that may affect the patient's understanding"
            (r'factors.*affect.*understanding', 'factors_hearing'),
            (r'ability.*cope.*hearing', 'factors_hearing'),

            # Adjustments for tribunal
            # T131/T133/T134: "Are there any adjustments that the tribunal may consider"
            (r'adjustments.*tribunal.*consider', 'adjustments'),
            (r'deal with the case fairly', 'adjustments'),

            # Strengths/positive factors
            # T131/T133/T134: "What are the strengths or positive factors"
            (r'strengths.*positive factors', 'strengths'),
            (r'what are the strengths', 'strengths'),

            # Progress/behaviour
            # T134: "current progress, engagement with nursing staff, behaviour"
            # T131: "current progress, behaviour, capacity and insight"
            # T133: "current progress, behaviour, compliance and insight"
            (r'current progress', 'progress'),
            (r'summary.*progress', 'progress'),
            (r'progress.*behaviour', 'progress'),
            (r'engagement with nursing', 'progress'),

            # Compliance/understanding of treatment
            # T134: "patient's understanding of, compliance with"
            # T131: "patient's understanding of, compliance with"
            (r'understanding.*compliance', 'compliance'),
            (r'compliance.*willingness', 'compliance'),
            (r'willingness to accept', 'compliance'),

            # Harm to self or others
            # T131/T133/T134: "incidents where the patient has harmed themselves or others"
            (r'harmed themselves or others', 'risk_harm'),
            (r'incidents.*harm', 'risk_harm'),
            (r'threatened to harm', 'risk_harm'),

            # Property damage
            # T131/T133/T134: "incidents where the patient has damaged property"
            (r'damaged property', 'risk_property'),
            (r'threatened to damage property', 'risk_property'),

            # Section 2 detention
            # T131/T133/T134: "In Section 2 cases, is detention in hospital justified"
            (r'section 2 cases.*detention', 's2_detention'),
            (r'in section 2 cases', 's2_detention'),

            # Other detention cases
            # T131/T133/T134: "In all other cases is the provision of medical treatment"
            (r'all other cases.*provision', 'other_detention'),
            (r'in all other cases', 'other_detention'),

            # Discharge risk
            # T131/T133/T134: "If the patient was discharged from hospital, would they be likely to act in a manner dangerous"
            (r'discharged.*dangerous', 'discharge_risk'),
            (r'discharged from hospital.*manner dangerous', 'discharge_risk'),
            (r'likely to act in a manner dangerous', 'discharge_risk'),

            # Community risk management
            # T131/T133/T134: "Please explain how risks could be managed effectively in the community"
            (r'risks.*managed.*community', 'community'),
            (r'managed effectively in the community', 'community'),
            (r'community.*conditions.*powers', 'community'),

            # Recommendations
            # T131/T133/T134: "Do you have any recommendations to the tribunal"
            (r'recommendations.*tribunal', 'recommendations'),
            (r'do you have any recommendations', 'recommendations'),

            # === T134 NURSING-SPECIFIC SECTIONS ===

            # Nursing care
            # T134: "What is the nature of nursing care and medication"
            (r'nature of nursing care', 'nursing_care'),
            (r'nursing care.*medication', 'nursing_care'),

            # Level of observation
            # T134: "To what level of observation is the patient currently subject"
            (r'level of observation', 'observation_level'),
            (r'observation.*subject', 'observation_level'),

            # Contact with relatives/friends/patients
            # T134: "Does the patient have contact with relatives, friends or other patients"
            (r'contact with relatives', 'contact'),
            (r'contact.*friends.*patients', 'contact'),
            (r'nature of that interaction', 'contact'),

            # Community support
            # T134: "What community support does the patient have"
            (r'what community support', 'community_support'),
            (r'community support.*have', 'community_support'),

            # AWOL
            # T134: "occasions on which the patient has been absent without leave"
            (r'absent without leave', 'awol'),
            (r'awol', 'awol'),

            # Seclusion/restraint
            # T134: "Have there been any occasions on which the patient has been secluded or restrained"
            (r'secluded or restrained', 'seclusion'),
            (r'seclusion.*restraint', 'seclusion'),
            (r'occasions.*seclusion', 'seclusion'),

            # Other relevant information
            # T133/T134: "Is there any other relevant information that the tribunal should know"
            (r'other relevant information', 'other_info'),
            (r'other.*information.*tribunal', 'other_info'),

            # === T131 MEDICAL-SPECIFIC (map to closest nursing equivalents) ===

            # Medical treatment → nursing_care (similar concept)
            # T131: "What appropriate and available medical treatment has been prescribed"
            (r'medical treatment.*prescribed', 'nursing_care'),
            (r'appropriate.*available.*treatment', 'nursing_care'),

            # === SKIP THESE (no good nursing equivalent) ===
            # - Index offence/forensic history
            # - Previous mental health dates
            # - Previous admission reasons
            # - Current admission circumstances
            # - Diagnosis
            # - Learning disability
            # - Detention requirement
            # - MCA/DoL
            # - Home/family circumstances
            # - Housing/accommodation
            # - Financial position
            # - Employment
            # - Care pathway/plan details
            # - Patient views
            # - Nearest relative views
            # - MAPPA
        ]

    def _match_heading_to_card(self, heading_text: str) -> str:
        """Match a heading to a nursing card key using regex patterns.

        Uses heading CONTENT to determine the section, not section numbers.
        This enables accurate cross-talk between different report types.
        """
        import re

        if not heading_text or len(heading_text) < 10:
            return None

        lower = heading_text.lower()
        patterns = self._get_heading_to_card_mapping()

        for pattern, card_key in patterns:
            if re.search(pattern, lower):
                return card_key

        return None

    def _parse_tribunal_docx(self, file_path: str) -> dict:
        """Parse a DOCX tribunal report and extract sections.

        Supports cross-talk: can import T131 (psychiatric), T134 (nursing),
        or social circumstances reports and map to nursing sections.
        """
        from docx import Document
        import re

        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"[NURSING] Failed to open DOCX: {e}")
            return {}

        # Get full text for detection
        full_text = ' '.join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += ' ' + cell.text

        # Detect report type
        report_type = self._detect_report_type(full_text)
        print(f"[NURSING] Detected report type: {report_type}")

        if report_type == 'unknown':
            print("[NURSING] Unknown report type - will try heading-based matching")

        result = {'form_type': report_type, 'sections': {}}

        # NOTE: We do NOT use section number mapping anymore.
        # All matching is done by HEADING CONTENT using _match_heading_to_card().
        # This ensures correct cross-talk between T131, T133, and T134 reports.

        # Helper to get unique cells (merged cells repeat same content)
        def get_unique_cells(cell_list):
            unique = []
            for c in cell_list:
                if c and (not unique or c != unique[-1]):
                    unique.append(c)
            return unique

        # Helper to detect if text is a question/heading (not actual content)
        def is_question_text(text):
            """Returns True if text appears to be a question heading, not answer content."""
            if not text or len(text) < 10:
                return False
            lower = text.lower().strip()

            # Question patterns - these are headings/questions, not answers
            question_patterns = [
                'are there any factors',
                'are there any adjustments',
                'what is the nature of',
                'to what level of observation',
                'does the patient have contact',
                'what community support',
                'what are the strengths',
                'give a summary of',
                'give details of any',
                'details of any occasions',
                'what is the patient',
                'have there been any occasions',
                'in section 2 cases',
                'in all other cases',
                'if the patient was discharged',
                'if the patient were discharged',  # variation
                'please explain how risks',
                'please explain how any risks',  # variation
                'is there any other relevant',
                'do you have any recommendations',
                'name of responsible clinician',
                'is the patient now suffering',
                'what appropriate and available',
                'what are the dates of',
                'what are the circumstances',
                'give reasons for any previous',
                'does the patient have a learning',
                'is there any mental disorder',
                'note: this report must be',
                'this report must be up-to-date',
                'full name',
                'date of birth',
                'usual place of residence',
                'patient details',
                'your name',
                'your role',
                'date of report',
                'would they be likely to act',  # discharge question
                'managed effectively in the community',  # community question
                'including the use of any lawful',  # community question
            ]

            for pattern in question_patterns:
                if pattern in lower:
                    return True

            # Also check for numbered section format: "1.", "2.", etc. at start
            import re
            if re.match(r'^\d{1,2}\.\s+', text):
                return True

            return False

        # Parse tables for section content
        for table in doc.tables:
            rows = list(table.rows)
            i = 0
            while i < len(rows):
                cells = [cell.text.strip() for cell in rows[i].cells]
                if not any(cells):
                    i += 1
                    continue

                first_cell = cells[0] if cells else ''
                unique_cells = get_unique_cells(cells)

                # Match section by HEADING CONTENT only (not section numbers)
                # This ensures correct cross-talk between T131, T133, T134 reports
                # Try ALL cells in the row - some formats have number in first cell, heading in second
                heading_key = None
                heading_cell_idx = -1
                for cell_idx, cell_text in enumerate(unique_cells):
                    # Skip pure section numbers like "2.", "3."
                    if re.match(r'^\d{1,2}\.\s*$', cell_text):
                        continue
                    matched = self._match_heading_to_card(cell_text)
                    if matched:
                        heading_key = matched
                        heading_cell_idx = cell_idx
                        break

                if heading_key:
                    # Check if answer is in next row or same row
                    answer = ""

                    # First check if answer is in cells AFTER the heading cell
                    start_idx = heading_cell_idx + 1 if heading_cell_idx >= 0 else 1
                    if len(unique_cells) > start_idx:
                        for cell_text in unique_cells[start_idx:]:
                            # Skip question/heading text - we want actual answers
                            if is_question_text(cell_text):
                                continue
                            if cell_text and not re.match(r'^\d+\.\s*', cell_text):
                                inline = self._parse_inline_checkbox(cell_text)
                                if inline:
                                    answer = inline
                                    break
                                cleaned = self._clean_docx_content(cell_text)
                                if cleaned and cleaned not in ('No', 'Yes', 'N/A'):
                                    answer = cleaned
                                    break

                    # If no answer in current row, check next row
                    if not answer and i + 1 < len(rows):
                        next_cells = [cell.text.strip() for cell in rows[i + 1].cells]
                        unique_next = get_unique_cells(next_cells)

                        for idx, cell_text in enumerate(unique_next):
                            # Skip section numbers and question text
                            if re.match(r'^\d+\.\s*', cell_text):
                                break
                            if is_question_text(cell_text):
                                continue

                            inline_answer = self._parse_inline_checkbox(cell_text)
                            if inline_answer:
                                answer = inline_answer
                                break

                            next_cell = unique_next[idx + 1] if idx + 1 < len(unique_next) else ""
                            yes_no_answer, _ = self._extract_yes_no_answer(cell_text, next_cell)
                            if yes_no_answer:
                                answer = yes_no_answer
                                break

                            if cell_text and not cell_text.startswith(('No\n', 'Yes\n')):
                                cleaned = self._clean_docx_content(cell_text)
                                # Only use if it's not question text
                                if cleaned and not is_question_text(cleaned):
                                    answer = cleaned
                                    break

                        if answer:
                            i += 1  # Skip the answer row

                    if answer:
                        # For Yes/No sections: if there's content, prefix with "Yes - "
                        yes_no_keys = {'factors_hearing', 'adjustments', 'contact', 'seclusion',
                                      'other_info', 'recommendations', 'discharge_risk', 's2_detention',
                                      'other_detention'}
                        if heading_key in yes_no_keys and answer not in ('Yes', 'No', 'N/A'):
                            answer = f"Yes - {answer}"
                        result['sections'][heading_key] = answer
                        print(f"[NURSING] Matched heading ({heading_key}): {answer[:50]}...")

                    i += 1
                    continue

                # No section match - try special patterns

                # Check for AWOL by heading text
                if any('absent without leave' in c.lower() for c in cells):
                    if i + 1 < len(rows):
                        next_cells = [cell.text.strip() for cell in rows[i + 1].cells]
                        unique_next = get_unique_cells(next_cells)
                        if unique_next:
                            answer = self._clean_docx_content(unique_next[0])
                            if answer and 'awol' not in result['sections']:
                                result['sections']['awol'] = answer
                                print(f"[NURSING] Section (awol): {answer[:50]}...")
                                i += 1

                # Check for patient details table
                elif 'Name of Patient' in first_cell or 'name of patient' in first_cell.lower():
                    if len(unique_cells) > 1:
                        result['sections']['patient_details'] = f"Name: {unique_cells[1]}"
                elif 'Date of Birth' in first_cell:
                    if len(unique_cells) > 1:
                        if 'patient_details' not in result['sections']:
                            result['sections']['patient_details'] = ""
                        result['sections']['patient_details'] += f"\nDOB: {unique_cells[1]}"
                elif 'NHS Number' in first_cell:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nNHS: {unique_cells[1]}"
                elif 'Usual Place of Residence' in first_cell:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nAddress: {unique_cells[1]}"
                elif 'Mental Health Act Status' in first_cell:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nMHA Status: {unique_cells[1]}"

                # Check for author info
                elif 'Your Name' in first_cell:
                    if len(unique_cells) > 1:
                        result['sections']['author'] = f"Name: {unique_cells[1]}"
                elif 'Your Role' in first_cell:
                    if len(unique_cells) > 1 and 'author' in result['sections']:
                        result['sections']['author'] += f"\nRole: {unique_cells[1]}"
                elif 'Date of Report' in first_cell:
                    if len(unique_cells) > 1 and 'author' in result['sections']:
                        result['sections']['author'] += f"\nDate: {unique_cells[1]}"

                # Check for signature table
                elif first_cell == 'Signed:':
                    if len(unique_cells) > 1:
                        result['sections']['signature'] = f"Signed: {unique_cells[1]}"
                elif first_cell == 'Print Name:' and 'signature' in result['sections']:
                    if len(unique_cells) > 1:
                        result['sections']['signature'] += f"\nName: {unique_cells[1]}"
                elif first_cell == 'Date:' and 'signature' in result['sections']:
                    if len(unique_cells) > 1:
                        result['sections']['signature'] += f"\nDate: {unique_cells[1]}"

                i += 1

        # ============================================================
        # PARAGRAPH PARSING - Extract checkbox selections and content
        # This handles the exported format where checkboxes are in paragraphs
        # ============================================================
        current_section = None
        current_section_key = None
        pending_checkbox = None  # Track if Yes checkbox was marked

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect section from numbered heading OR by question text
            section_match = re.match(r'^(\d{1,2})\.\s+', text)

            # Also check for section 2 which sometimes doesn't have "2." prefix
            is_section_2 = 'factors' in text.lower() and 'affect' in text.lower() and 'understanding' in text.lower()

            if section_match or is_section_2:
                # Save any pending checkbox selection
                if current_section_key and pending_checkbox and current_section_key not in result['sections']:
                    result['sections'][current_section_key] = pending_checkbox
                    print(f"[NURSING] Para section {current_section} ({current_section_key}): {pending_checkbox}")

                if is_section_2 and not section_match:
                    current_section = '2'
                    current_section_key = 'factors_hearing'
                else:
                    current_section = section_match.group(1)
                    # Use heading-based matching for paragraph content
                    current_section_key = self._match_heading_to_card(text)
                pending_checkbox = None
                continue

            # Detect checkbox selections
            if current_section_key:
                if text.startswith('☒'):
                    # Checked checkbox
                    if 'yes' in text.lower():
                        pending_checkbox = "Yes"
                        # Check if there's content after "Yes - "
                        yes_match = re.match(r'☒\s*Yes\s*[-–:]\s*(.+)', text, re.IGNORECASE)
                        if yes_match:
                            content_after = yes_match.group(1).strip()
                            # Only use if it's actual content, not question prompt
                            if content_after and not is_question_text(content_after):
                                pending_checkbox = f"Yes - {content_after}"
                    elif 'no' in text.lower():
                        pending_checkbox = "No"
                elif text.startswith('☐'):
                    # Unchecked checkbox - skip
                    pass
                elif pending_checkbox == "Yes" and not is_question_text(text) and len(text) > 10:
                    # This might be content following a Yes checkbox
                    if current_section_key not in result['sections']:
                        result['sections'][current_section_key] = f"Yes - {text}"
                        print(f"[NURSING] Para section {current_section} ({current_section_key}): Yes - {text[:50]}...")
                        pending_checkbox = None  # Consumed

        # Save final pending checkbox
        if current_section_key and pending_checkbox and current_section_key not in result['sections']:
            result['sections'][current_section_key] = pending_checkbox
            print(f"[NURSING] Para section {current_section} ({current_section_key}): {pending_checkbox}")

        print(f"[NURSING] Parsed DOCX: found {len(result['sections'])} sections")
        for key in result['sections']:
            print(f"  - {key}")
        return result

    def _clean_docx_content(self, text: str) -> str:
        """Clean up content from DOCX, removing checkbox formatting."""
        import re

        if not text:
            return ""

        # Remove checkbox patterns like [   ], [ X ], [x], etc.
        text = re.sub(r'\[\s*[xX]?\s*\]', '', text)

        # Remove "No" and "Yes" on their own lines
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            line = line.strip()
            if line.lower() in ('no', 'yes', 'n/a'):
                continue
            # Remove various "If yes" patterns at start of line
            line = re.sub(r'^[–-]\s*If yes[^:]*:\s*', '', line)
            line = re.sub(r'^If yes[^:]*:\s*', '', line)
            if line:
                cleaned_lines.append(line)

        result = '\n'.join(cleaned_lines).strip()
        return result

    def _parse_inline_checkbox(self, text: str) -> str:
        """Parse inline checkbox format like 'No [ ] Yes [X] N/A [ ]' to extract checked value."""
        import re

        if not text:
            return ""

        # Pattern for "Label [X]" or "Label [ ]" format
        # Look for checked boxes [x], [X], [ x], [x ], [ X ], etc.
        checked_pattern = r'(No|Yes|N/A)\s*\[\s*[xX]\s*\]'
        match = re.search(checked_pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)

        return ""

    def _extract_yes_no_answer(self, cell_text: str, next_cell: str = "") -> tuple:
        """Extract answer from Yes/No checkbox format. Returns (answer, has_explanation)."""
        import re

        # Check for inline format: "No [ ] Yes [X] N/A [ ]"
        inline_result = self._parse_inline_checkbox(cell_text)
        if inline_result:
            return inline_result, False

        # Check for two-cell format where first cell is "No\nYes" and second has checkboxes
        if cell_text.startswith(('No\n', 'Yes\n')) and next_cell:
            # Parse the checkbox cell - format is typically:
            # "[   ]\n[  X ] –  If yes, what are they?\n\nExplanation..."
            # First line is No checkbox, second line is Yes checkbox
            lines = next_cell.split('\n')

            # Find which checkbox is checked (has X)
            no_checked = False
            yes_checked = False
            for i, line in enumerate(lines):
                # Check if this line has a checked checkbox
                if re.search(r'\[\s*[xX]\s*\]', line):
                    if i == 0:  # First checkbox = No
                        no_checked = True
                    else:  # Second checkbox = Yes
                        yes_checked = True
                        break

            if no_checked and not yes_checked:
                return "No", False

            if yes_checked:
                # Look for explanation after "If yes"
                if 'If yes' in next_cell or 'if yes' in next_cell:
                    # Try pattern ending with colon, question mark, or end of line
                    match = re.search(r'[Ii]f yes[^\n]*\n+(.*)', next_cell, re.DOTALL)
                    if match:
                        explanation = match.group(1).strip()
                        if explanation:
                            return explanation, True
                    # Try splitting on double newline
                    parts = next_cell.split('\n\n')
                    if len(parts) > 1:
                        explanation = parts[1].strip()
                        if explanation:
                            return explanation, True
                else:
                    # No "If yes" - look for explanation directly after checkbox
                    # Format: "[ ]\n[ x ]  \nExplanation text..."
                    lines = next_cell.split('\n')
                    explanation_lines = []
                    found_checkbox = False
                    for line in lines:
                        if found_checkbox and line.strip():
                            explanation_lines.append(line.strip())
                        if re.search(r'\[\s*[xX]\s*\]', line):
                            found_checkbox = True
                    if explanation_lines:
                        explanation = '\n'.join(explanation_lines)
                        return explanation, True
                return "Yes", False

        # Check for checked box in the cell itself
        if '[x' in cell_text.lower():
            # Find what label is before the checked box
            lines = cell_text.split('\n')
            for i, line in enumerate(lines):
                if '[x' in line.lower() or '[ x' in line.lower():
                    # Check previous line or same line for label
                    if i > 0 and lines[i-1].strip().lower() in ('no', 'yes'):
                        return lines[i-1].strip(), False

        return "", False

    def _parse_scanned_pdf(self, file_path: str) -> dict:
        """Parse a scanned PDF using OCR and extract tribunal report sections."""
        import subprocess
        import tempfile
        import os
        import re

        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("[NURSING] PyMuPDF not available for OCR")
            return {}

        result = {'form_type': 'T134', 'sections': {}}

        try:
            doc = fitz.open(file_path)
        except Exception as e:
            print(f"[NURSING] Failed to open PDF: {e}")
            return {}

        # OCR all pages
        all_text = []
        for page in doc:
            pix = page.get_pixmap(dpi=200)
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                pix.save(f.name)
                temp_img = f.name

            try:
                ocr_result = subprocess.run(
                    ['tesseract', temp_img, 'stdout'],
                    capture_output=True, text=True, timeout=60
                )
                if ocr_result.returncode == 0:
                    all_text.append(ocr_result.stdout)
            except Exception as e:
                print(f"[NURSING] OCR error: {e}")
            finally:
                try:
                    os.unlink(temp_img)
                except:
                    pass

        doc.close()

        if not all_text:
            print("[NURSING] No text extracted from OCR")
            return {}

        full_text = '\n\n'.join(all_text)
        print(f"[NURSING] OCR extracted {len(full_text)} chars")

        # Check if this is a nursing report
        if 'nursing' not in full_text.lower() and 'nurse' not in full_text.lower():
            print("[NURSING] Not a nursing report")
            return {}

        # Section mapping
        section_map = {
            '2': 'factors_hearing',
            '3': 'adjustments',
            '4': 'nursing_care',
            '5': 'observation_level',
            '6': 'contact',
            '7': 'community_support',
            '8': 'strengths',
            '9': 'progress',
            '10': 'awol',
            '11': 'compliance',
            '12': 'risk_harm',
            '13': 'risk_property',
            '14': 'seclusion',
            '15': 's2_detention',
            '16': 'other_detention',
            '17': 'discharge_risk',
            '18': 'community',
            '19': 'recommendations',
        }

        # Yes/No sections - these have checkbox answers
        yes_no_sections = {'2', '3', '6', '10', '12', '14', '15', '16', '17', '19'}

        print(f"[NURSING] Extracting sections using question-based approach...")

        # Define the FULL question text for each section (more reliable than section numbers)
        # Each entry: (section_num, key, question_pattern, is_yes_no)
        section_questions = [
            ('2', 'factors_hearing', r'factors.*?(?:affect|hearing).*?\?', True),
            ('3', 'adjustments', r'adjustments.*?tribunal.*?(?:fairly|justly).*?\?', True),
            ('4', 'nursing_care', r'nature of nursing care.*?\?', False),
            ('5', 'observation_level', r'level of observation.*?subject.*?\?', False),
            ('6', 'contact', r'contact with.*?(?:relatives|friends|patients).*?\?', True),
            ('7', 'community_support', r'community support.*?(?:have|available).*?\?', False),
            ('8', 'strengths', r'strengths.*?positive factors.*?\?', False),
            ('9', 'progress', r'(?:summary|progress).*?(?:behaviour|cooperation|insight).*?\?', False),
            ('10', 'awol', r'absent without leave.*?\?', False),
            ('11', 'compliance', r'compliance.*?willingness.*?medication.*?\?', False),
            ('12', 'risk_harm', r'harmed themselves or others.*?\?', False),
            ('13', 'risk_property', r'damaged property.*?\?', False),
            ('14', 'seclusion', r'secluded or restrained.*?\?', True),
            ('15', 's2_detention', r'Section 2 cases.*?protection of others.*?\?', True),
            ('16', 'other_detention', r'all other cases.*?protection of others.*?\?', True),
            ('17', 'discharge_risk', r'discharged.*?dangerous.*?(?:themselves|others).*?\?', True),
            ('18', 'community', r'(?:managed|risks).*?community.*?(?:conditions|powers).*?\?', False),
            ('19', 'recommendations', r'recommendations.*?tribunal.*?\?', True),
        ]

        # Find each question in the text and extract the answer that follows
        for sec_num, key, q_pattern, is_yes_no in section_questions:
            # Find the question
            q_match = re.search(q_pattern, full_text, re.IGNORECASE | re.DOTALL)
            if not q_match:
                print(f"[NURSING] Section {sec_num} ({key}): Question not found")
                continue

            q_end = q_match.end()
            print(f"[NURSING] Section {sec_num} question found at {q_match.start()}-{q_end}")

            # Find the next section's question to determine where this answer ends
            next_q_start = len(full_text)
            for next_sec, next_key, next_pattern, _ in section_questions:
                if int(next_sec) > int(sec_num):
                    next_match = re.search(next_pattern, full_text, re.IGNORECASE | re.DOTALL)
                    if next_match and next_match.start() > q_end:
                        if next_match.start() < next_q_start:
                            next_q_start = next_match.start()

            # Also check for Signature as end marker
            sig_match = re.search(r'\bSignature\b', full_text[q_end:], re.IGNORECASE)
            if sig_match:
                sig_pos = q_end + sig_match.start()
                if sig_pos < next_q_start:
                    next_q_start = sig_pos

            # Extract content between question end and next question start
            content = full_text[q_end:next_q_start].strip()

            # For Yes/No questions, also look for the "If yes" follow-up
            if is_yes_no:
                # Check for "If yes" pattern and skip to after it
                if_yes_match = re.search(r'If yes[^?]*\?', content, re.IGNORECASE)
                if if_yes_match:
                    content = content[if_yes_match.end():].strip()

            print(f"[NURSING] Section {sec_num} raw ({len(content)} chars): {content[:60]}...")

            # Clean the content
            cleaned = self._clean_ocr_section(content, sec_num, yes_no_sections)

            if cleaned:
                result['sections'][key] = cleaned
                print(f"[NURSING] Section {sec_num} ({key}): {cleaned[:60]}...")

        # Try to extract patient details from section 1
        # Look for patient name - use more specific patterns
        patient_patterns = [
            # Look for "Full name of the patient" followed by actual name
            r'Full name of the patient[:\s]*\n*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)',
            # Look for "Mr/Mrs/Ms" followed by name
            r'\b(Mr\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'\b(Mrs\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
            r'\b(Ms\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
        ]
        false_positives = ['hospital', 'logo', 'report', 'nursing', 'tribunal', 'ward', 'tamarind', 'signature']
        for pat in patient_patterns:
            patient_match = re.search(pat, full_text, re.IGNORECASE)
            if patient_match:
                name = patient_match.group(1).strip()
                # Filter out false positives
                name_lower = name.lower()
                if not any(fp in name_lower for fp in false_positives):
                    result['sections']['patient_details'] = f"Name: {name}"
                    print(f"[NURSING] Found patient name: {name}")
                    break

        print(f"[NURSING] OCR parsed {len(result['sections'])} sections")
        return result

    def _clean_ocr_section(self, text: str, section_num: str, yes_no_sections: set) -> str:
        """Clean up OCR'd content for a section, handling Yes/No questions properly."""
        import re

        if not text:
            return ""

        # IMPORTANT: Detect checkbox states BEFORE any text cleanup
        # This must happen first because cleanup removes checkbox markers
        original_text = text  # Keep original for checkbox detection

        # Check for checked/unchecked No box - ASCII and Unicode
        # ASCII: [X] No, [x] No, [v] No = checked; [_] No, [ ] No = unchecked
        # Unicode: ☒ No, ☑ No = checked; ☐ No = unchecked
        no_checked = bool(re.search(r'(?:\[[xXv]\]|☒|☑)\s*No\b', original_text))
        no_unchecked = bool(re.search(r'(?:\[[\s_]\]|☐)\s*No\b', original_text))

        # Check for checked/unchecked Yes box
        yes_checked = bool(re.search(r'(?:\[[xXv]\]|☒|☑)\s*Yes\b', original_text))
        yes_unchecked = bool(re.search(r'(?:\[[\s_]\]|☐)\s*Yes\b', original_text))

        # Remove common OCR artifacts (now safe since we've captured checkbox state)
        # Remove all checkbox markers - ASCII and Unicode
        text = re.sub(r'\[[\s_xXv]?\]', '', text)  # [_], [X], [x], [v], [ ], []
        text = re.sub(r'[☐☒☑]', '', text)  # Unicode checkboxes
        text = re.sub(r'Page\s*\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'Crown copyright.*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'T134\b.*?(?:\n|$)', '', text)
        text = re.sub(r'T7134\b.*?(?:\n|$)', '', text)
        text = re.sub(r'7134\b.*?(?:\n|$)', '', text)
        text = re.sub(r'In-Patient.*?Report.*?(?:\n|$)', '', text, flags=re.IGNORECASE)

        # Remove section numbers that appear on their own lines (OCR artifacts)
        text = re.sub(r'^\s*\d{1,2}\.\s*$', '', text, flags=re.MULTILINE)

        # Question patterns to remove from content (including partial fragments)
        question_patterns = [
            r'Are there any factors.*?hearing\??',
            r'Does the patient.*?adjustments.*?\??',
            r'What is the nature of nursing care.*?\??',
            r'To what level of observation.*?\??',
            r'Does the patient have.*?contact with relatives.*?\??',
            r'What community support.*?\??',
            r'What are the strengths.*?\??',
            r'Give a summary.*?progress.*?\??',
            r'Details of any.*?absent without leave.*?\??',
            r'What is the patient.*?compliance.*?medication\??',
            r'Give details.*?harmed themselves or others\??',
            r'Give details.*?damaged property\??',
            r'Have there been any occasions.*?secluded or restrained\??',
            r'In Section 2 cases.*?protection of others\??',
            r'In all other cases.*?protection of others\??',
            r'If the patient was discharged.*?dangerous.*?\??',
            r'Please explain how.*?managed in the community\??',
            r'Do you have any recommendations.*?tribunal\??',
            r'Please provide your recommendations.*?below\.?',
            r'If yes.*?below\.?',
            r'What is the nature of that interaction\??',
            r'What are they\??',
            # Partial question fragments that might appear
            r',?\s*or threatened harm to others\.?',
            r',?\s*or threatened to damage\s*property\.?',
            r',?\s*or threatened to damage',
            r'including the use of any',
            r'lawful conditions or recall powers',
            r'engagement with nursing staff.*?insight',
            r'cooperation.*?activities.*?self-care',
        ]

        for pattern in question_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.DOTALL)

        # Remove leading dashes and artifacts
        text = re.sub(r'^[\s\-_,\.]+', '', text)
        text = re.sub(r'\n[\s\-_]+\n', '\n', text)

        # Clean up multiple newlines and spaces
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'  +', ' ', text)
        text = text.strip()

        # Handle Yes/No sections
        if section_num in yes_no_sections:
            # Use checkbox states detected at start of function (before cleanup removed markers)
            # Variables no_checked, no_unchecked, yes_checked, yes_unchecked were set earlier

            lines = text.strip().split('\n')

            no_selected = False
            yes_selected = False
            content_lines = []
            found_answer = False

            # Determine selection based on checkbox states
            if no_checked and not yes_checked:
                no_selected = True
                found_answer = True
            elif yes_checked and not no_checked:
                yes_selected = True
                found_answer = True
            elif yes_unchecked and not yes_checked:
                # Yes checkbox is unchecked = No is selected
                no_selected = True
                found_answer = True

            # Check for "Yes - [content]" pattern which indicates Yes is selected with content
            # This handles cases like "[_] No\n\nYes - Please provide your recommendations..."
            yes_content_match = re.search(r'Yes\s*-\s*(.+)', text, re.IGNORECASE | re.DOTALL)
            if yes_content_match and not found_answer:
                potential_content = yes_content_match.group(1).strip()
                # Make sure it's not just question text - filter out question prompts
                question_starters = ('what are', 'please explain', 'please provide', 'if yes',
                                    'do you have', 'are there', 'does the', 'have there')
                if potential_content and not potential_content.lower().startswith(question_starters):
                    # Check if there's actual content (not just form headers)
                    content_lines_check = [l.strip() for l in potential_content.split('\n')
                                          if l.strip() and not any(x in l.lower() for x in ['7134', 't134', 'nursing report'])]
                    if content_lines_check:
                        yes_selected = True
                        found_answer = True
                        no_selected = False  # Override any No detection
                else:
                    # The "Yes - " was followed by question text, look for actual content after it
                    # Find where the actual answer content starts (after question prompt)
                    lines_after_yes = potential_content.split('\n')
                    actual_content = []
                    for line in lines_after_yes:
                        line = line.strip()
                        if not line:
                            continue
                        # Skip lines that are question prompts
                        if any(q in line.lower() for q in ['please provide', 'please explain', 'what are',
                                                           'if yes', 'do you have', 'recommendations below',
                                                           '7134', 't134', 'nursing report']):
                            continue
                        actual_content.append(line)
                    if actual_content:
                        yes_selected = True
                        found_answer = True
                        no_selected = False

            # Only check for standalone "No" if we haven't found Yes with content
            if not found_answer:
                first_word = text.strip().split()[0] if text.strip() else ""
                if first_word == 'No':
                    no_selected = True
                    found_answer = True

            # Parse lines for content
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Skip question text lines
                if any(q in line.lower() for q in ['are there any', 'does the patient', 'have there been',
                                                    'in section 2', 'in all other', 'if the patient',
                                                    'do you have any', 'please provide', 'what are they',
                                                    'what is the nature', 'if yes', 'in-patient',
                                                    'nursing report', '7134', 't134', 'please explain']):
                    continue

                # Skip standalone checkbox labels
                if line in ('No', 'Yes', 'N/A', 'None'):
                    if not found_answer:
                        # Standalone No/Yes without checkbox detection
                        if line == 'No':
                            no_selected = True
                            found_answer = True
                        elif line == 'Yes':
                            yes_selected = True
                            found_answer = True
                    continue

                # Handle "Yes - content" pattern
                if line.startswith('Yes -'):
                    rest = line[5:].strip()
                    if rest.lower().startswith(('what are', 'please', 'if yes')):
                        continue  # Question text, skip
                    yes_selected = True
                    found_answer = True
                    if rest:
                        content_lines.append(rest)
                    continue

                # Handle "No - content" pattern (rare)
                if line.startswith('No -'):
                    no_selected = True
                    found_answer = True
                    rest = line[4:].strip()
                    if rest:
                        content_lines.append(rest)
                    continue

                # Only collect content if Yes is selected (not No)
                if not no_selected:
                    if yes_selected or (not found_answer and line not in ('No', 'Yes')):
                        # Skip form headers
                        if not any(q in line.lower() for q in ['7134', 't134', 'nursing report']):
                            content_lines.append(line)
                            if not found_answer:
                                yes_selected = True
                                found_answer = True

            content_after = '\n'.join(content_lines).strip()

            # Return appropriate value - No takes priority if selected
            if no_selected:
                return "No"
            elif yes_selected and content_after:
                # Add Yes prefix for sections that need it
                if section_num in ('2', '3', '6', '19'):
                    return f"Yes - {content_after}"
                return content_after
            elif content_after:
                # Has content but no explicit Yes/No found
                if section_num in ('2', '3', '6', '19'):
                    return f"Yes - {content_after}"
                return content_after
            else:
                return "None"

        else:
            # Non Yes/No sections - just clean up the content
            lines = text.split('\n')
            cleaned = []
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Skip standalone Yes/No/None
                if line in ('No', 'Yes', 'N/A'):
                    continue
                cleaned.append(line)

            result = '\n'.join(cleaned).strip()

            # Handle empty or "None" responses
            if not result or result.lower() == 'none':
                return "None"

            return result

    def _populate_from_docx(self, result: dict, file_path: str):
        """Populate report sections from parsed DOCX tribunal form data.

        Supports cross-talk: can import from T131 (psychiatric), T134 (nursing),
        or social circumstances reports. Matching sections are mapped by heading.
        """
        from PySide6.QtWidgets import QMessageBox
        import os

        sections = result.get('sections', {})
        form_type = result.get('form_type', 'unknown')

        # Map report types to friendly names
        type_names = {
            'T131': 'Psychiatric/Medical (T131)',
            'T134': 'Nursing (T134)',
            'social': 'Social Circumstances',
            'unknown': 'Tribunal Report'
        }
        source_name = type_names.get(form_type, form_type)

        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        filename = os.path.basename(file_path)
        action = self._ask_import_action(filename, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        # Populate the cards
        loaded_count = 0
        loaded_sections = []
        merged_sections = {}
        for key, content in sections.items():
            if key in self.cards and content:
                card = self.cards[key]
                if hasattr(card, 'editor'):
                    if action == "add":
                        content = self._merge_report_section(key, content, filename)
                    self._imported_report_data[key] = content
                    merged_sections[key] = content
                    card.editor.setPlainText(content)
                    loaded_count += 1
                    loaded_sections.append(key)
                    print(f"[NURSING] Loaded section: {key}")

        # Push patient details to shared store for other forms
        self._push_patient_details_to_shared_store(sections)

        # Push sections to shared store for cross-talk with psychiatric form
        from shared_data_store import get_shared_store
        shared_store = get_shared_store()
        shared_store.set_report_sections(merged_sections, source_form="nursing_tribunal")

        # Show success message with cross-talk info
        cross_talk_note = ""
        if form_type != 'T134' and form_type != 'unknown':
            cross_talk_note = f"\n(Cross-talk import from {source_name})"

        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"Successfully {action_word.lower()} from:\n{filename}\n\n"
            f"Source: {source_name}{cross_talk_note}\n\n"
            f"{action_word} {loaded_count} sections.\n\n"
            f"Review and edit the content as needed."
        )

        print(f"[NURSING] {action_word} {loaded_count} sections from {source_name}")

        # Populate popups with imported data (Yes/No states, details, collapsible sections)
        self._populate_popups_with_imported_data(merged_sections)

    def _push_patient_details_to_shared_store(self, sections: dict):
        """Extract patient details from loaded sections and push to shared store."""
        import re
        from shared_data_store import get_shared_store

        patient_details = sections.get('patient_details', '')
        if not patient_details:
            return

        patient_info = {
            "name": None,
            "dob": None,
            "nhs_number": None,
            "gender": None,
            "age": None,
            "ethnicity": None,
            "address": None,
            "mha_status": None,
        }

        from datetime import datetime

        # Extract name
        name_match = re.search(r"(?:Name|Patient)[:\s]+([A-Za-z][A-Za-z\-\' ]+)", patient_details, re.IGNORECASE)
        if name_match:
            patient_info["name"] = name_match.group(1).strip()

        # Extract DOB - try numeric format first
        dob_match = re.search(r"(?:DOB|Date of Birth|D\.O\.B)[:\s]+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})", patient_details, re.IGNORECASE)
        if dob_match:
            dob_str = dob_match.group(1).strip()
            for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y"]:
                try:
                    patient_info["dob"] = datetime.strptime(dob_str, fmt)
                    break
                except ValueError:
                    continue

        # Try text date format: "7 October 1979", "15 Jan 1985", etc.
        if not patient_info["dob"]:
            text_dob_match = re.search(
                r"(?:DOB|Date of Birth|D\.O\.B)[:\s]+(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{4})",
                patient_details, re.IGNORECASE
            )
            if text_dob_match:
                day = int(text_dob_match.group(1))
                month_str = text_dob_match.group(2)
                year = int(text_dob_match.group(3))
                month_map = {
                    'january': 1, 'jan': 1, 'february': 2, 'feb': 2, 'march': 3, 'mar': 3,
                    'april': 4, 'apr': 4, 'may': 5, 'june': 6, 'jun': 6,
                    'july': 7, 'jul': 7, 'august': 8, 'aug': 8, 'september': 9, 'sep': 9,
                    'october': 10, 'oct': 10, 'november': 11, 'nov': 11, 'december': 12, 'dec': 12
                }
                month = month_map.get(month_str.lower(), 1)
                try:
                    patient_info["dob"] = datetime(year, month, day)
                except ValueError:
                    pass

        # Extract address
        address_match = re.search(r"(?:Address|Residence|Place of Residence)[:\s]+(.+?)(?:\n|MHA|NHS|$)", patient_details, re.IGNORECASE)
        if address_match:
            patient_info["address"] = address_match.group(1).strip()

        # Extract MHA status
        mha_match = re.search(r"(?:MHA Status|Mental Health Act|Section)[:\s]+([^\n]+)", patient_details, re.IGNORECASE)
        if mha_match:
            patient_info["mha_status"] = mha_match.group(1).strip()

        # Extract NHS number
        nhs_match = re.search(r"(?:NHS)[:\s]*(\d{3}\s*\d{3}\s*\d{4}|\d{10})", patient_details, re.IGNORECASE)
        if nhs_match:
            nhs = nhs_match.group(1).replace(" ", "")
            patient_info["nhs_number"] = f"{nhs[:3]} {nhs[3:6]} {nhs[6:]}" if len(nhs) == 10 else nhs

        # Extract gender
        gender_match = re.search(r"(?:Gender|Sex)[:\s]*(Male|Female|M|F)\b", patient_details, re.IGNORECASE)
        if gender_match:
            g = gender_match.group(1).upper()
            patient_info["gender"] = "Male" if g in ("MALE", "M") else "Female" if g in ("FEMALE", "F") else None

        # Extract age
        age_patterns = [
            r"(?:AGE)[:\s]*(\d{1,3})\s*(?:years?|yrs?|y\.?o\.?)?\b",
            r"\b(\d{1,3})\s*(?:year|yr)\s*old\b",
            r"\b(\d{1,3})\s*y\.?o\.?\b",
            r"\baged?\s*(\d{1,3})\b",
        ]
        for pattern in age_patterns:
            match = re.search(pattern, patient_details, re.IGNORECASE)
            if match:
                age_val = int(match.group(1))
                if 0 < age_val < 120:
                    patient_info["age"] = age_val
                    break

        # Calculate age from DOB if not found explicitly
        if not patient_info["age"] and patient_info["dob"]:
            today = datetime.today()
            dob = patient_info["dob"]
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if 0 < age < 120:
                patient_info["age"] = age

        # Extract ethnicity
        ethnicity_patterns = [
            r"(?:ETHNICITY|ETHNIC\s*(?:GROUP|ORIGIN)?)[:\s]+([A-Za-z][A-Za-z\s\-\/]+?)(?:\n|$|,)",
            r"\b(White\s*(?:British|Irish|European|Other)?)\b",
            r"\b(Black\s*(?:British|African|Caribbean|Other)?)\b",
            r"\b(Asian\s*(?:British|Indian|Pakistani|Bangladeshi|Chinese|Other)?)\b",
            r"\b(Mixed\s*(?:White\s*(?:and|&)\s*(?:Black\s*(?:Caribbean|African)|Asian))?)\b",
        ]
        for pattern in ethnicity_patterns:
            match = re.search(pattern, patient_details, re.IGNORECASE)
            if match:
                ethnicity_val = match.group(1).strip()
                if len(ethnicity_val) > 2:
                    patient_info["ethnicity"] = ethnicity_val.title()
                    break

        # Push to shared store if any info found
        if any(patient_info.values()):
            shared_store = get_shared_store()
            shared_store.set_patient_info(patient_info, source="nursing_tribunal")
            print(f"[NURSING] Pushed patient details to SharedDataStore: {list(k for k,v in patient_info.items() if v)}")

    def _send_to_data_extractor(self, file_path: str):
        """Send a file to the data extractor for processing."""
        self._data_extractor_source_file = file_path
        self._open_data_extractor_overlay()

        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            if hasattr(self._data_extractor_overlay, 'load_file'):
                self._data_extractor_overlay.load_file(file_path)
            elif hasattr(self._data_extractor_overlay, 'upload_and_extract'):
                self._data_extractor_overlay.upload_and_extract()
            print(f"[NURSING] Sent to data extractor: {file_path}")

    def _ask_import_action(self, source_filename: str, import_type: str = "report") -> str:
        """Ask user whether to add to existing imported data or replace it.

        Returns: 'add', 'replace', or 'cancel'
        """
        from PySide6.QtWidgets import QMessageBox

        if import_type == "report":
            has_existing = hasattr(self, '_imported_report_data') and any(self._imported_report_data.values())
        else:
            has_existing = hasattr(self, '_extracted_raw_notes') and bool(self._extracted_raw_notes)

        if not has_existing:
            return "replace"

        msg = QMessageBox(self)
        msg.setWindowTitle("Import Data")
        if import_type == "report":
            msg.setText(
                f"Report data has already been imported.\n\n"
                f"New file: {source_filename}\n\n"
                f"Would you like to add this data to the existing import, or replace it?"
            )
        else:
            msg.setText(
                f"Clinical notes have already been loaded.\n\n"
                f"Would you like to add these notes to the existing set, or replace them?"
            )
        add_btn = msg.addButton("Add to Existing", QMessageBox.AcceptRole)
        replace_btn = msg.addButton("Replace All", QMessageBox.DestructiveRole)
        cancel_btn = msg.addButton("Cancel", QMessageBox.RejectRole)
        msg.setDefaultButton(add_btn)
        msg.exec()

        clicked = msg.clickedButton()
        if clicked == cancel_btn:
            return "cancel"
        elif clicked == replace_btn:
            return "replace"
        return "add"

    def _merge_report_section(self, key: str, new_content: str, source_filename: str) -> str:
        """Merge new report content with existing section, adding source reference."""
        existing = self._imported_report_data.get(key, "")
        if not existing:
            return new_content
        return f"{existing}\n\n─── Source: {source_filename} ───\n\n{new_content}"

    def _populate_from_pdf(self, result: dict, file_path: str):
        """Populate report sections from parsed PDF tribunal form data."""
        from PySide6.QtWidgets import QMessageBox
        from pdf_loader import format_radio_value
        import os

        form_type = result.get('form_type', 'unknown')
        sections = result.get('sections', {})

        # Map section keys to card keys for nursing report
        # T134 nursing report has different section structure
        section_to_card = {
            'patient_details': 'patient_details',
            'author': 'author',
            'factors_hearing': 'factors_hearing',
            'adjustments': 'adjustments',
            'forensic': 'forensic',
            'previous_mh_dates': 'previous_mh',
            'diagnosis': 'diagnosis',
            'treatment': 'treatment',
            'strengths': 'strengths',
            'progress': 'progress',
            'compliance': 'compliance',
            'risk_harm': 'risk_harm',
            'risk_property': 'risk_property',
            'discharge_risk': 'discharge_risk',
            'community': 'community',
            'recommendations': 'recommendations',
            'signature_date': 'signature',
        }

        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        filename = os.path.basename(file_path)
        action = self._ask_import_action(filename, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        # Populate the cards
        loaded_count = 0
        merged_sections = {}
        for section_key, content in sections.items():
            card_key = section_to_card.get(section_key)
            if card_key and card_key in self.cards:
                card = self.cards[card_key]
                if hasattr(card, 'editor'):
                    # Format radio button values
                    if section_key in ['learning_disability', 'detention_required', 's2_detention', 'other_detention']:
                        content = format_radio_value(content)
                    if action == "add":
                        content = self._merge_report_section(card_key, content, filename)
                    self._imported_report_data[card_key] = content
                    merged_sections[card_key] = content
                    card.editor.setPlainText(content)
                    loaded_count += 1
                    print(f"[PDF] Loaded section: {section_key}")

        # Push patient details to shared store for other forms
        self._push_patient_details_to_shared_store(sections)

        # Show success message
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "PDF Loaded",
            f"Successfully {action_word.lower()} {form_type} report from:\n{filename}\n\n"
            f"{action_word} {loaded_count} sections from the PDF.\n\n"
            f"Review and edit the content as needed."
        )

        print(f"[PDF] {action_word} {loaded_count} sections from {form_type} report")

        # Populate popups with imported data collapsible sections
        self._populate_popups_with_imported_data(merged_sections)

    def _populate_popups_with_imported_data(self, sections: dict):
        """Populate popups with imported data - creates popups and calls _populate_single_popup."""
        for section_key, content in sections.items():
            if not content or not content.strip():
                continue

            # Create popup if it doesn't exist (popups are normally created lazily)
            if section_key not in self.popups:
                popup = self._create_popup(section_key)
                if popup:
                    self.popups[section_key] = popup
                    self.popup_stack.addWidget(popup)
                    if hasattr(popup, 'sent'):
                        popup.sent.connect(lambda text, k=section_key: self._update_card(k, text))
                    print(f"[NURSING] Created popup '{section_key}' during import")

            if section_key in self.popups:
                popup = self.popups[section_key]
                self._populate_single_popup(popup, section_key, content)

    def _populate_single_popup(self, popup, section_key: str, content: str):
        """Populate a single popup with imported content - copied from psychiatric tribunal."""
        import re

        # ============================================================
        # HELPER FUNCTIONS FOR CHECKBOX DETECTION
        # ============================================================
        def has_yes_cross(text):
            """Check if Yes has a checked box symbol."""
            if re.search(r'yes\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☒☑✓✔\[xX\]].*no\s*[☐□\[\s\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes.*[☐□]\s*no', text, re.IGNORECASE):
                return True
            return False

        def has_no_cross(text):
            """Check if No has a checked box symbol (meaning answer is No)."""
            if re.search(r'no\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            if re.search(r'no\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*no\b', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☐□\[\s\]].*no\s*[☒☑✓✔\[xX\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☐□]\s*yes.*[☒☑✓✔]\s*no', text, re.IGNORECASE):
                return True
            return False

        def is_yes_content(text):
            """Check if content indicates Yes answer."""
            if has_yes_cross(text):
                return True
            if has_no_cross(text):
                return False
            lower = text.lower().strip()
            if lower.startswith("yes") or "yes -" in lower or "yes," in lower:
                return True
            # If substantial text and not explicitly "no", assume yes
            if len(text.strip()) > 30 and not lower.startswith("no"):
                return True
            return False

        def extract_detail(text):
            """Extract detail text after Yes/No prefix."""
            for prefix in ["Yes - ", "Yes, ", "Yes-", "yes - ", "yes, ", "yes-", "Yes ", "yes "]:
                if text.startswith(prefix):
                    return text[len(prefix):].strip()
            return text

        # ============================================================
        # PATIENT DETAILS - parse text into fields directly
        # ============================================================
        if section_key == "patient_details":
            from datetime import datetime
            patient_info = {}

            name_match = re.search(r"(?:Full\s*Name|Name|Patient)[:\s]+(.+)", content, re.IGNORECASE)
            if name_match:
                patient_info["name"] = name_match.group(1).strip()

            dob_match = re.search(r"(?:Date of Birth|DOB|D\.O\.B)[:\s]+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})", content, re.IGNORECASE)
            if dob_match:
                dob_str = dob_match.group(1).strip()
                for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y"]:
                    try:
                        patient_info["dob"] = datetime.strptime(dob_str, fmt)
                        break
                    except ValueError:
                        continue

            if not patient_info.get("dob"):
                text_dob = re.search(
                    r"(?:Date of Birth|DOB|D\.O\.B)[:\s]+(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{4})",
                    content, re.IGNORECASE
                )
                if text_dob:
                    month_map = {'january':1,'jan':1,'february':2,'feb':2,'march':3,'mar':3,'april':4,'apr':4,'may':5,'june':6,'jun':6,'july':7,'jul':7,'august':8,'aug':8,'september':9,'sep':9,'october':10,'oct':10,'november':11,'nov':11,'december':12,'dec':12}
                    try:
                        patient_info["dob"] = datetime(int(text_dob.group(3)), month_map.get(text_dob.group(2).lower(), 1), int(text_dob.group(1)))
                    except ValueError:
                        pass

            gender_match = re.search(r"(?:Gender|Sex)[:\s]*(Male|Female|M|F)\b", content, re.IGNORECASE)
            if gender_match:
                g = gender_match.group(1).upper()
                patient_info["gender"] = "Male" if g in ("MALE", "M") else "Female"

            address_match = re.search(r"(?:Usual Place of Residence|Address|Residence)[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
            if address_match:
                patient_info["address"] = address_match.group(1).strip()

            if patient_info and hasattr(popup, 'fill_patient_info'):
                popup.fill_patient_info(patient_info)
                print(f"[NURSING] Filled patient_details fields: {list(patient_info.keys())}")
            return

        # ============================================================
        # YES/NO POPUPS WITH STANDARD yes_btn/no_btn
        # ============================================================
        yes_no_sections = ["factors_hearing", "adjustments", "compliance",
                          "discharge_risk", "recommendations"]

        if section_key in yes_no_sections:
            is_yes = is_yes_content(content)
            print(f"[NURSING] Section '{section_key}' content starts: {repr(content[:100])}")
            print(f"[NURSING] Section '{section_key}' is_yes={is_yes}")

            if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
                if is_yes:
                    popup.yes_btn.setChecked(True)
                    print(f"[NURSING] Set '{section_key}' popup to Yes")
                else:
                    popup.no_btn.setChecked(True)
                    print(f"[NURSING] Set '{section_key}' popup to No")

                # For factors_hearing, also detect and set the specific factor radio buttons
                if section_key == "factors_hearing" and is_yes:
                    content_lower = content.lower()
                    if hasattr(popup, 'autism_rb') and ('autism' in content_lower or 'autistic' in content_lower):
                        popup.autism_rb.setChecked(True)
                        print(f"[NURSING] Set factors_hearing to Autism")
                    elif hasattr(popup, 'ld_rb') and ('learning disability' in content_lower or 'learning difficulties' in content_lower):
                        popup.ld_rb.setChecked(True)
                        print(f"[NURSING] Set factors_hearing to Learning Disability")
                    elif hasattr(popup, 'patience_rb') and ('irritab' in content_lower or 'frustration' in content_lower or 'patience' in content_lower):
                        popup.patience_rb.setChecked(True)
                        print(f"[NURSING] Set factors_hearing to Low frustration tolerance")

                # Populate the details_field (pre-existing) with imported text
                detail = extract_detail(content)
                if detail and detail.lower() not in ("yes", "no"):
                    if hasattr(popup, 'details_field'):
                        popup.details_field.setPlainText(detail)
                        print(f"[NURSING] Set '{section_key}' details: {detail[:50]}...")
                    elif hasattr(popup, 'additional_details_field'):
                        popup.additional_details_field.setPlainText(detail)
                        print(f"[NURSING] Set '{section_key}' additional details: {detail[:50]}...")

                # Hide the always-visible additional details for factors_hearing/adjustments
                if section_key in ("factors_hearing", "adjustments"):
                    if hasattr(popup, 'always_visible_details'):
                        popup.always_visible_details.hide()

                # Send updated text to card
                if hasattr(popup, '_send_to_card'):
                    popup._send_to_card()
                    print(f"[NURSING] Sent '{section_key}' popup text to card")

        # ============================================================
        # NURSING CARE SECTION (section 4) - Parse medication list
        # ============================================================
        if section_key == "nursing_care" and hasattr(popup, '_medications'):
            self._populate_medications_from_import(popup, content)

        # Add collapsible imported data section
        self._add_imported_data_to_popup(popup, section_key, content)

    def _add_imported_data_to_popup(self, popup, section_key: str, content: str):
        """Add imported data collapsible section to popup with checkboxes."""
        from PySide6.QtWidgets import QLabel, QVBoxLayout, QWidget, QCheckBox, QHBoxLayout, QFrame
        from PySide6.QtCore import Qt
        import re

        # Skip if empty or minimal checkbox-only content
        if not content or not content.strip():
            return
        cleaned = content.strip()
        if re.match(r'^(No|Yes)\s*\[\s*\]\s*(No|Yes)\s*\[\s*\]$', cleaned, re.IGNORECASE):
            return
        if re.match(r'^[☐☒\s]*(No|Yes)\s*[☐☒\s]*(No|Yes)\s*[☐☒\s]*$', cleaned, re.IGNORECASE):
            return
        if cleaned.lower() in ('yes', 'no', 'n/a', 'yes.', 'no.'):
            return

        # Skip heading-only sections (they don't need imported data)
        skip_sections = {'s2_detention', 'other_detention', 'author', 'signature', 'patient_details', 'factors_hearing', 'adjustments'}
        if section_key in skip_sections:
            return

        # Skip if popup already has imported data
        if getattr(popup, '_imported_data_added', False):
            return

        # Try to find target layout - more robust approach
        target_layout = None
        parent_widget = None

        if hasattr(popup, 'scroll_layout') and popup.scroll_layout:
            target_layout = popup.scroll_layout
            parent_widget = target_layout.parentWidget()
        elif hasattr(popup, 'main_layout') and popup.main_layout:
            target_layout = popup.main_layout
            parent_widget = target_layout.parentWidget()
        elif hasattr(popup, 'container_layout') and popup.container_layout:
            target_layout = popup.container_layout
            parent_widget = target_layout.parentWidget()
        elif hasattr(popup, 'layout'):
            layout_attr = popup.layout
            if hasattr(layout_attr, 'insertWidget'):
                target_layout = layout_attr
                parent_widget = target_layout.parentWidget()

        if not target_layout or not callable(getattr(target_layout, 'insertWidget', None)):
            print(f"[NURSING] No target layout found for '{section_key}'")
            return

        if not parent_widget:
            parent_widget = popup

        try:
            from background_history_popup import CollapsibleSection

            # Create collapsible section
            import_section = CollapsibleSection("Imported Data", parent=parent_widget, start_collapsed=False)
            import_section.set_content_height(150)
            import_section._min_height = 80
            import_section._max_height = 2000
            import_section.set_header_style("""
                QFrame {
                    background: rgba(220, 252, 231, 0.95);
                    border: 1px solid rgba(20, 184, 166, 0.5);
                    border-radius: 6px 6px 0 0;
                }
            """)
            import_section.title_label.setStyleSheet("""
                QLabel {
                    font-size: 17px;
                    font-weight: 600;
                    color: #0d9488;
                    background: transparent;
                    border: none;
                }
            """)

            # Create content widget
            content_widget = QWidget()
            content_widget.setStyleSheet("""
                QWidget {
                    background: rgba(220, 252, 231, 0.95);
                    border: 1px solid rgba(20, 184, 166, 0.4);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            content_layout = QVBoxLayout(content_widget)
            content_layout.setContentsMargins(12, 10, 12, 10)
            content_layout.setSpacing(4)

            # Get current card content to check what's already there
            card_text = ""
            if section_key in self.cards:
                card = self.cards[section_key]
                if hasattr(card, 'editor'):
                    card_text = card.editor.toPlainText().lower()

            # Split content into paragraphs and create checkbox + label pairs
            popup._imported_checkboxes = []

            # Sections where all content should be treated as ONE paragraph
            single_paragraph_sections = {'recommendations', 'progress', 'risk_harm', 'risk_property', 'community'}

            if section_key in single_paragraph_sections:
                # Use a scrollable QTextEdit instead of QLabel for long content
                from PySide6.QtWidgets import QTextEdit as _QTextEdit
                cleaned_content = content.strip()

                # Scale height based on content length (~60 chars per line at 13px)
                est_lines = max(8, len(cleaned_content) // 55 + 2)
                text_height = min(600, est_lines * 18)  # ~18px per line
                section_height = text_height + 40

                import_section.set_content_height(section_height)
                import_section._max_height = 2000

                cb = QCheckBox()
                cb.setFixedSize(20, 20)
                cb.setStyleSheet("QCheckBox::indicator { width: 16px; height: 16px; }")
                cb.setProperty("full_text", cleaned_content)
                cb.setProperty("section_key", section_key)

                cb.setChecked(False)
                cb.toggled.connect(lambda checked, sk=section_key, txt=cleaned_content: self._on_imported_checkbox_toggled(sk, txt, checked))

                row = QHBoxLayout()
                row.setContentsMargins(0, 4, 0, 4)
                row.setSpacing(8)
                row.setAlignment(Qt.AlignmentFlag.AlignTop)
                row.addWidget(cb, 0, Qt.AlignmentFlag.AlignTop)

                text_edit = _QTextEdit()
                text_edit.setPlainText(cleaned_content)
                text_edit.setReadOnly(True)
                text_edit.setMinimumHeight(text_height)
                text_edit.setStyleSheet("""
                    QTextEdit {
                        font-size: 17px; color: #4a4a4a;
                        background: transparent; border: none;
                    }
                """)
                row.addWidget(text_edit, 1)

                row_widget = QFrame()
                row_widget.setStyleSheet("QFrame { background: transparent; border: none; }")
                row_widget.setLayout(row)
                content_layout.addWidget(row_widget)

                popup._imported_checkboxes = [cb]

                import_section.set_content(content_widget)
                target_layout.insertWidget(0, import_section)
                # Push content to top (absorb extra space from widgetResizable scroll area)
                target_layout.addStretch()
                popup._imported_data_added = True
                print(f"[NURSING] Added single-paragraph imported data for '{section_key}'")
                return

            elif '\n\n' in content:
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            else:
                paragraphs = [p.strip() for p in content.split('\n') if p.strip()]

            if paragraphs:
                for para in paragraphs:
                    # Create container for checkbox + label
                    item_container = QFrame()
                    item_container.setStyleSheet("QFrame { background: transparent; border: none; }")
                    item_layout = QHBoxLayout(item_container)
                    item_layout.setContentsMargins(0, 4, 0, 4)
                    item_layout.setSpacing(8)
                    item_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

                    # Small checkbox (no text)
                    cb = QCheckBox()
                    cb.setFixedSize(20, 20)
                    cb.setStyleSheet("QCheckBox::indicator { width: 16px; height: 16px; }")
                    cb.setProperty("full_text", para)
                    cb.setProperty("section_key", section_key)

                    cb.setChecked(False)

                    # Connect to handler
                    cb.toggled.connect(lambda checked, sk=section_key, txt=para: self._on_imported_checkbox_toggled(sk, txt, checked))

                    item_layout.addWidget(cb, 0, Qt.AlignmentFlag.AlignTop)

                    # Word-wrapped label for the text
                    text_label = QLabel(para)
                    text_label.setWordWrap(True)
                    text_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
                    text_label.setStyleSheet("font-size: 17px; color: #4a4a4a; background: transparent; padding: 0;")
                    text_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
                    item_layout.addWidget(text_label, 1)

                    popup._imported_checkboxes.append(cb)
                    content_layout.addWidget(item_container)
            else:
                # Fallback: just add as selectable label if no paragraphs
                content_label = QLabel(content)
                content_label.setWordWrap(True)
                content_label.setStyleSheet("font-size: 17px; color: #4a4a4a; background: transparent; border: none;")
                content_layout.addWidget(content_label)

            import_section.set_content(content_widget)

            # Insert at top of popup
            target_layout.insertWidget(0, import_section)
            # Push content to top (absorb extra space from widgetResizable scroll area)
            target_layout.addStretch()
            popup._imported_data_added = True
            print(f"[NURSING] Added imported data with {len(paragraphs)} checkboxes to popup '{section_key}'")

        except Exception as e:
            print(f"[NURSING] Error adding imported data to '{section_key}': {e}")

    def _on_imported_checkbox_toggled(self, section_key: str, text: str, checked: bool):
        """Handle when an imported data checkbox is toggled."""
        if section_key not in self.cards:
            return

        card = self.cards[section_key]
        if not hasattr(card, 'editor'):
            return

        current_text = card.editor.toPlainText()

        if checked:
            # Add text to card if not already there
            if text not in current_text:
                if current_text.strip():
                    new_text = current_text.strip() + "\n\n" + text
                else:
                    new_text = text
                card.editor.setPlainText(new_text)
                print(f"[NURSING] Added imported text to '{section_key}'")
        else:
            # Remove text from card
            if text in current_text:
                new_text = current_text.replace(text, "").strip()
                # Clean up double newlines
                while "\n\n\n" in new_text:
                    new_text = new_text.replace("\n\n\n", "\n\n")
                card.editor.setPlainText(new_text)
                print(f"[NURSING] Removed imported text from '{section_key}'")

    def _populate_medications_from_import(self, popup, content: str):
        """Parse medication list from imported content and populate popup medication entries."""
        import re
        from CANONICAL_MEDS import MEDICATIONS

        if not content:
            return

        print(f"[NURSING] Parsing medications from imported content...")

        # Frequency mapping
        FREQ_MAP = {
            "od": "OD", "once daily": "OD", "daily": "OD", "mane": "OD", "nocte": "Nocte",
            "bd": "BD", "twice daily": "BD", "twice a day": "BD",
            "tds": "TDS", "three times daily": "TDS", "three times a day": "TDS",
            "qds": "QDS", "four times daily": "QDS", "four times a day": "QDS",
            "prn": "PRN", "as required": "PRN", "when required": "PRN",
            "weekly": "Weekly", "once weekly": "Weekly", "1 weekly": "Weekly",
            "fortnightly": "Fortnightly", "2 weekly": "Fortnightly", "every 2 weeks": "Fortnightly",
            "3 weekly": "3 Weekly", "every 3 weeks": "3 Weekly",
            "monthly": "Monthly", "4 weekly": "Monthly", "every 4 weeks": "Monthly",
        }

        # Parse medication lines - look for patterns like:
        # - Olanzapine 10mg OD
        # • Depakote 500mg BD
        # Sertraline 100mg once daily
        lines = content.split('\n')
        meds_found = []

        # Build a lowercase -> uppercase key map for matching
        med_name_map = {}
        for med_key, med_info in MEDICATIONS.items():
            med_name_map[med_key.lower()] = med_key
            canonical = med_info.get("canonical", "").lower()
            if canonical:
                med_name_map[canonical] = med_key
            for pattern in med_info.get("patterns", []):
                med_name_map[pattern.lower()] = med_key

        for line in lines:
            line = line.strip()
            # Remove bullet points and list markers
            line = re.sub(r'^[\-\•\*\d\.]+\s*', '', line)
            if not line:
                continue

            # Try to match medication name
            words = line.split()
            med_key = None
            med_name = None

            # Check first 1-3 words for medication name
            for i in range(1, min(4, len(words) + 1)):
                test_name = ' '.join(words[:i]).lower()
                # Remove trailing punctuation
                test_name = re.sub(r'[,.:;]+$', '', test_name)
                if test_name in med_name_map:
                    med_key = med_name_map[test_name]
                    med_name = ' '.join(words[:i])
                    break

            if not med_key:
                # Try matching against patterns in the full line
                line_lower = line.lower()
                for name, key in med_name_map.items():
                    if re.search(r'\b' + re.escape(name) + r'\b', line_lower):
                        med_key = key
                        med_name = name
                        break

            if not med_key:
                continue

            # Extract dose (number followed by mg, mcg, g, ml, etc.)
            dose_match = re.search(r'(\d+(?:\.\d+)?)\s*(mg|mcg|g|ml|micrograms?)', line, re.IGNORECASE)
            dose = dose_match.group(0) if dose_match else ""

            # Extract frequency
            freq = "OD"  # Default
            line_lower = line.lower()
            for freq_pattern, freq_value in FREQ_MAP.items():
                if freq_pattern in line_lower:
                    freq = freq_value
                    break

            meds_found.append({
                "med_key": med_key,
                "dose": dose,
                "freq": freq
            })
            print(f"[NURSING] Found medication: {med_key} {dose} {freq}")

        if not meds_found:
            print(f"[NURSING] No medications found in imported content")
            return

        # Clear existing empty medication entries (keep only the first one if empty)
        while len(popup._medications) > 1:
            entry = popup._medications[-1]
            if entry["name"].currentText().strip() == "":
                popup._medications.remove(entry)
                entry["widget"].deleteLater()
            else:
                break

        # Add enough entries
        while len(popup._medications) < len(meds_found):
            popup._add_medication_entry()

        # Populate the medication entries
        for i, med in enumerate(meds_found):
            if i >= len(popup._medications):
                break

            entry = popup._medications[i]
            med_key = med.get("med_key", "")

            # Set medication name
            name_combo = entry.get("name")
            if name_combo and med_key:
                idx = name_combo.findText(med_key)
                if idx >= 0:
                    name_combo.setCurrentIndex(idx)
                else:
                    # Set as text (editable combo)
                    name_combo.setCurrentText(med_key)

            # Set dose
            dose_combo = entry.get("dose")
            if dose_combo and med.get("dose"):
                dose_str = med.get("dose")
                idx = dose_combo.findText(dose_str)
                if idx >= 0:
                    dose_combo.setCurrentIndex(idx)
                else:
                    dose_combo.setCurrentText(dose_str)

            # Set frequency
            freq_combo = entry.get("freq")
            if freq_combo and med.get("freq"):
                freq = med.get("freq")
                idx = freq_combo.findText(freq)
                if idx >= 0:
                    freq_combo.setCurrentIndex(idx)

        print(f"[NURSING] ✓ Pre-filled {len(meds_found)} medication(s) from import")

    def _open_data_extractor_overlay(self):
        """Create the data extractor (hidden) for background processing."""
        from data_extractor_popup import DataExtractorPopup

        if not hasattr(self, '_data_extractor_overlay') or not self._data_extractor_overlay:
            self._data_extractor_overlay = DataExtractorPopup(parent=self)
            self._data_extractor_overlay.hide()

            if hasattr(self._data_extractor_overlay, 'data_extracted'):
                self._data_extractor_overlay.data_extracted.connect(self._on_data_extracted)

    # Mapping from PTR card keys to NTR card keys for report imports
    # NTR1=PTR1, NTR2=PTR3, NTR3=PTR4, NTR4=PTR12, NTR8=PTR13,
    # NTR9=PTR14, NTR11=PTR15, NTR12=PTR17, NTR13=PTR18,
    # NTR17=PTR21, NTR18=PTR22
    PTR_TO_NTR_MAP = {
        "patient_details": "patient_details",   # PTR1  -> NTR1
        "factors_hearing": "factors_hearing",   # PTR3  -> NTR2
        "adjustments": "adjustments",           # PTR4  -> NTR3
        "treatment": "nursing_care",             # PTR12 -> NTR4
        "strengths": "strengths",                # PTR13 -> NTR8
        "progress": "progress",                  # PTR14 -> NTR9
        "compliance": "compliance",              # PTR15 -> NTR11
        "risk_harm": "risk_harm",                # PTR17 -> NTR12
        "risk_property": "risk_property",        # PTR18 -> NTR13
        "discharge_risk": "discharge_risk",      # PTR21 -> NTR17
        "community": "community",                # PTR22 -> NTR18
    }

    def _on_data_extracted(self, data: dict):
        """Handle extracted data from the data extractor and populate fixed panels."""
        import os
        print(f"[NURSING] Data extracted: {list(data.keys())}")
        cov = data.get("_coverage")
        if cov and cov.get("uncategorised", 0) > 0:
            print(f"[NURSING] Warning: {cov['uncategorised']} paragraphs uncategorised "
                  f"({cov['categorised']}/{cov['total_paragraphs']} categorised)")

        # Check if this is filtered data to send to the current card
        filtered_category = data.get("filtered_category")
        if filtered_category and hasattr(self, '_selected_card_key') and self._selected_card_key:
            print(f"[NURSING] Filtered category '{filtered_category}' -> sending to current card '{self._selected_card_key}'")
            self._send_filtered_to_current_card(data)
            return

        # Skip if this exact data was already processed
        categories = data.get("categories", {})
        cat_keys = tuple(sorted(categories.keys())) if categories else ()
        cat_count = sum(len(v.get("items", [])) if isinstance(v, dict) else 0 for v in categories.values())
        content_sig = (cat_keys, cat_count)
        if self._data_processed_id == content_sig:
            print(f"[NURSING] Skipping _on_data_extracted - data already processed")
            return
        self._data_processed_id = content_sig

        # Check if data came from a report (not notes)
        is_report = False
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            dtype = getattr(self._data_extractor_overlay, '_detected_document_type', None)
            if dtype == "reports":
                is_report = True
                print(f"[NURSING] Detected report data (dtype={dtype})")

        categories = data.get("categories", {})
        print(f"[NURSING] Available categories: {list(categories.keys())}")

        if is_report and categories:
            # Report pipeline: map PTR categories to NTR sections
            source = os.path.basename(getattr(self, '_data_extractor_source_file', '') or '') or "Data Extractor"
            self._populate_from_report_categories(categories, source_filename=source)
        else:
            # Notes pipeline - skip if report data already imported (prevents cross-talk)
            if self._has_report_data():
                print(f"[NURSING] Skipping notes pipeline - report data already imported")
                return

            raw_notes = []
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                raw_notes = getattr(self._data_extractor_overlay, 'notes', [])

            # Fall back to SharedDataStore if no local notes
            if not raw_notes:
                try:
                    from shared_data_store import get_shared_store
                    shared_store = get_shared_store()
                    if shared_store.has_notes():
                        raw_notes = shared_store.notes
                        print(f"[NURSING] Got {len(raw_notes)} notes from SharedDataStore (global import)")
                except Exception as e:
                    print(f"[NURSING] Error getting notes from SharedDataStore: {e}")

            print(f"[NURSING] Raw notes available: {len(raw_notes)}")

            # Ask add/replace if existing notes
            action = self._ask_import_action("", "notes")
            if action == "cancel":
                return
            if action == "add":
                existing_notes = getattr(self, '_extracted_raw_notes', []) or []
                raw_notes = existing_notes + raw_notes
                existing_cats = getattr(self, '_extracted_categories', {}) or {}
                for cat_name, cat_data in categories.items():
                    if cat_name in existing_cats and isinstance(existing_cats[cat_name], dict) and isinstance(cat_data, dict):
                        existing_items = existing_cats[cat_name].get("items", [])
                        new_items = cat_data.get("items", [])
                        existing_cats[cat_name]["items"] = existing_items + new_items
                    else:
                        existing_cats[cat_name] = cat_data
                categories = existing_cats

            # STORE at page level for popups created later
            self._extracted_raw_notes = raw_notes
            self._extracted_categories = categories
            self._incident_data = raw_notes
            print(f"[NURSING] Stored {len(raw_notes)} raw notes and {len(categories)} categories at page level")

            # Populate any existing popups
            self._populate_fixed_panels()

    def _populate_from_report_categories(self, categories: dict, source_filename: str = ""):
        """Populate NTR sections from PTR report categories.

        Maps PTR card keys to NTR card keys using PTR_TO_NTR_MAP and
        populates card editors + popups with imported data.
        """
        from PySide6.QtWidgets import QMessageBox
        from shared_data_store import get_shared_store

        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}
        if not hasattr(self, '_imported_report_sections'):
            self._imported_report_sections = {}

        source_label = source_filename or "Data Extractor"
        action = self._ask_import_action(source_label, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()
            self._imported_report_sections.clear()

        # Valid NTR card keys
        valid_ntr_keys = {key for _, key in self.SECTIONS}

        # Build NTR card_key -> combined text AND original PTR keys for cross-talk
        card_texts = {}
        all_ptr_texts = {}  # Original PTR keys for SharedDataStore cross-talk

        for cat_name, cat_data in categories.items():
            # Combine all items' text first (needed for both NTR and cross-talk)
            items = cat_data.get("items", []) if isinstance(cat_data, dict) else []
            texts = []
            for item in items:
                if isinstance(item, dict):
                    text = item.get("text", "") or item.get("content", "")
                elif isinstance(item, str):
                    text = item
                else:
                    text = str(item)
                if text:
                    texts.append(text.strip())

            if not texts:
                continue

            combined = "\n\n".join(texts)

            # Always store under original PTR key for cross-talk
            if cat_name in all_ptr_texts:
                all_ptr_texts[cat_name] += "\n\n" + combined
            else:
                all_ptr_texts[cat_name] = combined

            # Map PTR card key to NTR card key for local population
            if cat_name in self.PTR_TO_NTR_MAP:
                ntr_key = self.PTR_TO_NTR_MAP[cat_name]
            elif cat_name in valid_ntr_keys:
                ntr_key = cat_name
            else:
                print(f"[NURSING] No NTR mapping for PTR category: {cat_name} (will pass through to other forms)")
                continue

            if ntr_key in card_texts:
                card_texts[ntr_key] += "\n\n" + combined
            else:
                card_texts[ntr_key] = combined

            # Track sections
            if ntr_key not in self._imported_report_sections:
                self._imported_report_sections[ntr_key] = []
            self._imported_report_sections[ntr_key].append((cat_name, combined))

        # Store imported data and populate NTR popups
        sections_for_store = {}
        for ntr_key, content in card_texts.items():
            if action == "add":
                content = self._merge_report_section(ntr_key, content, source_label)
            self._imported_report_data[ntr_key] = content
            sections_for_store[ntr_key] = content
            print(f"[NURSING] Stored report data for NTR section '{ntr_key}'")

        # Populate popups with imported data (flag to skip notes-based searches)
        self._is_report_import = True
        self._populate_popups_with_imported_data(sections_for_store)
        self._is_report_import = False

        # Push patient details to shared store
        self._push_patient_details_to_shared_store(sections_for_store)

        # Push ALL original PTR keys to shared store for cross-talk
        # This ensures PTR/SCT get sections that don't map to NTR (forensic, diagnosis, etc.)
        shared_store = get_shared_store()
        cross_talk_sections = dict(all_ptr_texts)
        # Also include NTR-mapped sections under their NTR keys for forms that share keys
        cross_talk_sections.update(sections_for_store)
        shared_store.set_report_sections(cross_talk_sections, source_form="nursing_tribunal")

        # Show success
        mapped_count = len(sections_for_store)
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"{action_word} report data from {source_label}.\n\n"
            f"{action_word} {mapped_count} PTR sections to NTR sections.\n\n"
            f"Click each card to review and send the content."
        )
        print(f"[NURSING] {action_word} {mapped_count} PTR sections to NTR sections")

    def _send_filtered_to_current_card(self, data: dict):
        """Send filtered data from data extractor to the currently selected card."""
        current_key = self._selected_card_key
        if not current_key:
            print("[NURSING] No card selected - cannot send filtered data")
            return

        categories = data.get("categories", {})
        filtered_category = data.get("filtered_category", "")

        # Combine all items from the filtered categories
        all_items = []
        for cat_name, cat_data in categories.items():
            if isinstance(cat_data, dict) and "items" in cat_data:
                all_items.extend(cat_data.get("items", []))

        if not all_items:
            print(f"[NURSING] No items found in filtered data for '{filtered_category}'")
            return

        # Format the text content
        texts = []
        for item in all_items:
            text = item.get("text", "").strip()
            date_str = item.get("date", "")
            if text:
                if date_str:
                    texts.append(f"[{date_str}] {text}")
                else:
                    texts.append(text)

        combined_text = "\n\n".join(texts)
        print(f"[NURSING] Sending {len(texts)} items to card '{current_key}'")

        # Update the card with the combined text
        if current_key in self.cards:
            self._update_card(current_key, combined_text)
            print(f"[NURSING] Card '{current_key}' updated with filtered data")

    # Category to section mapping for nursing report
    CATEGORY_SECTION_MAP = {
        "medication history": ["nursing_care"],
        "background history": ["contact", "community_support"],
        "history of presenting complaint": ["progress"],
        "risk": ["awol", "risk_harm", "risk_property", "seclusion", "other_detention", "discharge_risk"],
        "plan": ["community", "compliance"],
        "summary": ["other_info", "recommendations"],
    }

    def set_notes(self, notes: list):
        """
        Set notes from shared data store.
        Called by MainWindow when notes are available from another section.
        """
        if not notes:
            return

        # Skip if report data exists (report takes priority over notes)
        if self._has_report_data():
            print(f"[Nursing] Skipping set_notes - report data already imported")
            return

        # Skip if these exact notes were already processed
        notes_sig = (len(notes), id(notes))
        if self._notes_processed_id == notes_sig:
            print(f"[Nursing] Skipping set_notes - notes already processed")
            return
        self._notes_processed_id = notes_sig

        # Store raw notes at page level for use in sections
        self._extracted_raw_notes = notes

        # If data extractor exists, update its notes too
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            if hasattr(self._data_extractor_overlay, 'set_notes'):
                self._data_extractor_overlay.set_notes(notes)

        print(f"[Nursing] Received {len(notes)} notes from shared store")

    def _populate_fixed_panels(self):
        """Populate all fixed panels with extracted data."""
        import re
        from datetime import datetime
        from pathlib import Path

        raw_notes = getattr(self, '_extracted_raw_notes', [])
        categories = getattr(self, '_extracted_categories', {})

        if not raw_notes:
            print("[NURSING] No extracted data to populate panels")
            return

        # Helper to get items from category (categories have "items" key)
        def get_category_items(cat_name):
            for key, cat in categories.items():
                key_str = str(key) if key is not None else ""
                if cat_name.lower() in key_str.lower():
                    if isinstance(cat, dict):
                        return cat.get("items", [])
            return []

        # Section 7: Community Support - Use Social History or Personal History from data extractor
        if "community_support" in self.popups:
            popup = self.popups["community_support"]
            if hasattr(popup, 'set_entries'):
                # Try Social History first, then Background History (personal history), then fall back to raw notes
                entries = get_category_items("Social History")
                if not entries:
                    entries = get_category_items("Personal History")
                if not entries:
                    entries = raw_notes[:30] if raw_notes else []
                popup.set_entries(entries, f"{len(entries)} notes")
                print(f"[NURSING] Populated section 7 (community_support) with {len(entries)} entries")

        # Section 9: Progress - Use notes from last 6 months
        # Calculate notes first (before if block so it's available for _pending_section_data)
        from datetime import timedelta

        def parse_note_date_s9(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # First, find the most recent note date
        all_note_dates = []
        for n in raw_notes:
            note_date = parse_note_date_s9(n.get('date') or n.get('datetime'))
            if note_date:
                all_note_dates.append(note_date)

        if all_note_dates:
            most_recent_date = max(all_note_dates)
            oldest_date = min(all_note_dates)
            # 6 months before the most recent entry
            six_months_cutoff = most_recent_date - timedelta(days=180)

            print(f"[NURSING] Section 9 DEBUG: Most recent note: {most_recent_date.strftime('%d/%m/%Y')}")
            print(f"[NURSING] Section 9 DEBUG: Oldest note: {oldest_date.strftime('%d/%m/%Y')}")
            print(f"[NURSING] Section 9 DEBUG: 6-month cutoff: {six_months_cutoff.strftime('%d/%m/%Y')}")

            # Filter notes from last 6 months (relative to most recent entry)
            notes_with_dates = []
            for n in raw_notes:
                note_date = parse_note_date_s9(n.get('date') or n.get('datetime'))
                if note_date and note_date >= six_months_cutoff:
                    notes_with_dates.append(n)

            # Sort by date (most recent first)
            sorted_notes_s9 = sorted(
                notes_with_dates,
                key=lambda x: parse_note_date_s9(x.get('date') or x.get('datetime')),
                reverse=True
            )
            recent_progress = sorted_notes_s9

            # Show actual date range of filtered notes
            if recent_progress:
                filtered_dates = [parse_note_date_s9(n.get('date') or n.get('datetime')) for n in recent_progress]
                filtered_dates = [d for d in filtered_dates if d]
                if filtered_dates:
                    print(f"[NURSING] Section 9 DEBUG: Filtered range: {min(filtered_dates).strftime('%d/%m/%Y')} to {max(filtered_dates).strftime('%d/%m/%Y')}")
        else:
            # Fallback: no parseable dates, use first 100 raw notes
            recent_progress = raw_notes[:100]

        print(f"[NURSING] Section 9: Prepared {len(recent_progress)} notes from last 6 months (total raw: {len(raw_notes)})")

        if "progress" in self.popups:
            popup = self.popups["progress"]
            if hasattr(popup, 'notes'):
                popup.notes = raw_notes

            if hasattr(popup, 'set_entries'):
                popup.set_entries(recent_progress, f"{len(recent_progress)} notes (last 6 months)")
                print(f"[NURSING] Populated section 9 (progress) with {len(recent_progress)} notes from last 6 months")

        # Section 10: AWOL - search all notes
        if "awol" in self.popups:
            popup = self.popups["awol"]
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[NURSING] AWOL popup searching {len(raw_notes)} notes")

        # ================================================================
        # INCIDENT EXTRACTION for sections 12, 13, 14 (like psychiatric report sections 17, 18)
        # ================================================================

        # Load incident terms from dictionary file
        incident_terms = []
        incident_file = Path(__file__).parent / "incidentDICT.txt"
        if incident_file.exists():
            with open(incident_file, 'r', encoding='utf-8') as f:
                for line in f:
                    term = line.strip().lower()
                    if term:
                        incident_terms.append(term)
            print(f"[NURSING] Loaded {len(incident_terms)} incident terms from incidentDICT.txt")

        exclude_keywords = ['h/o', 'history of', 'previous noted', 'previously noted', 'previous history',
                            'risk of', 'historical risk', 'past risk',
                            'no aggression', 'nil aggression', 'nil aggressive', 'not aggressive',
                            'no aggressive', 'low risk and nil', 'no risk', 'nil on the shift',
                            'fire safety', 'fire test', 'gas / fire', 'gas/fire',
                            'risk and aggression:', 'risk/aggression', 'risk & aggression',
                            'politely declined', 'hearing 30th', 'managers hearing', 'manangers hearing',
                            'solicitor today', 'no signs of aggression', 'no physical or verbal aggression',
                            'did not display any aggressive', 'no aggressive behaviour',
                            'no challenging or aggressive', 'no episode of aggressive',
                            'no irritable or aggressive', 'no physical/verbal aggression',
                            'no instances of irritability, aggressive',
                            'did not appear intoxicated', 'not intoxicated', 'no intoxication',
                            'although did not appear intox', 'did not appear to be intox',
                            'if no further abusive', 'no further abusive']

        def is_sexual_health_or_history(text):
            """Check if 'sexual' is followed by 'health', 'history', 'hx', or 'transmitted' within 2 words."""
            # Match 'sexual' followed by 0-2 words then 'health', 'history', 'hx'
            pattern1 = r'\bsexual\s+(?:\w+\s+)?(?:health|history|hx)\b'
            # Match 'sexually transmitted'
            pattern2 = r'\bsexually\s+transmitted\b'
            return bool(re.search(pattern1, text, re.IGNORECASE) or re.search(pattern2, text, re.IGNORECASE))

        def parse_note_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # Extract individual incident lines
        incident_lines = []
        seclusion_lines = []  # Separate list for seclusion entries
        word_boundary_terms = {'anger'}

        for note in raw_notes:
            content = note.get('content', '') or note.get('text', '') or ''
            content = content.replace('\r\n', '\n').replace('\r', '\n')

            date_val = note.get('date') or note.get('datetime')
            date_obj = parse_note_date(date_val)
            date_str = date_obj.strftime('%d/%m/%Y') if date_obj else 'Unknown'

            # Split by newlines first, then by sentences
            all_segments = []
            for line in content.split('\n'):
                sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', line)
                all_segments.extend(sentences)

            for idx, line in enumerate(all_segments):
                line_clean = line.strip()
                line_clean = ''.join(c for c in line_clean if ord(c) >= 32 or c == '\t')
                line_lower = line_clean.lower()

                if not line_clean or len(line_clean) < 10:
                    continue

                # Skip if contains exclusion keywords
                if any(ex in line_lower for ex in exclude_keywords):
                    continue

                # Skip if "sexual" followed by "health" or "history" (not actual incidents)
                if is_sexual_health_or_history(line_lower):
                    continue

                # Check for seclusion keyword (for section 14)
                if 'seclusion' in line_lower:
                    seclusion_lines.append({
                        'date': date_str,
                        'date_obj': date_obj,
                        'text': f"{date_str}: {line_clean}",
                        'content': f"{date_str}: {line_clean}",
                        'term': 'seclusion'
                    })

                # Check for incident terms (for sections 12, 13)
                for term in incident_terms:
                    if term in word_boundary_terms:
                        pattern = r'\b' + re.escape(term) + r'\b'
                        if not re.search(pattern, line_lower):
                            continue
                    else:
                        if term not in line_lower:
                            continue

                    # For DATIX or Ulysses, include 2-3 more lines after
                    final_text = line_clean
                    if 'datix' in line_lower or 'ulysses' in line_lower:
                        extra_lines = []
                        for extra_idx in range(1, 4):
                            if idx + extra_idx < len(all_segments):
                                extra_line = all_segments[idx + extra_idx].strip()
                                extra_line = ''.join(c for c in extra_line if ord(c) >= 32 or c == '\t')
                                if extra_line and len(extra_line) > 5:
                                    extra_lines.append(extra_line)
                        if extra_lines:
                            final_text = line_clean + ' | ' + ' | '.join(extra_lines)

                    incident_lines.append({
                        'date': date_str,
                        'date_obj': date_obj,
                        'text': f"{date_str}: {final_text}",
                        'content': f"{date_str}: {final_text}",
                        'term': term
                    })
                    break  # Only one match per line

        # Sort incidents by date (most recent first)
        incident_lines.sort(key=lambda x: x['date_obj'] or datetime.min, reverse=True)
        seclusion_lines.sort(key=lambda x: x['date_obj'] or datetime.min, reverse=True)

        # Deduplicate by line content
        seen_lines = set()
        incidents = []
        for inc in incident_lines:
            line_key = inc['text'].lower()
            if line_key not in seen_lines:
                seen_lines.add(line_key)
                incidents.append(inc)

        seen_seclusion = set()
        seclusion_incidents = []
        for inc in seclusion_lines:
            line_key = inc['text'].lower()
            if line_key not in seen_seclusion:
                seen_seclusion.add(line_key)
                seclusion_incidents.append(inc)

        print(f"[NURSING] Found {len(incidents)} unique incident lines, {len(seclusion_incidents)} seclusion entries")

        # Store incident data for later use
        self._incident_data = incidents

        # Section 12: Risk Harm - use TribunalRiskHarmPopup with raw notes
        if "risk_harm" in self.popups:
            popup = self.popups["risk_harm"]
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[NURSING] Populated section 12 (risk_harm) with {len(raw_notes)} notes for harm search")

        # Section 13: Risk Property - use TribunalRiskPropertyPopup with raw notes
        if "risk_property" in self.popups:
            popup = self.popups["risk_property"]
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[NURSING] Populated section 13 (risk_property) with {len(raw_notes)} notes for property search")

        # Section 14: Seclusion/Restraint - use TribunalSeclusionPopup with raw notes
        if "seclusion" in self.popups:
            popup = self.popups["seclusion"]
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[NURSING] Populated section 14 (seclusion) with {len(raw_notes)} notes for seclusion search")

        # Section 17: Discharge Risk - use GPRRiskPopup with notes for risk analysis (same as psych tribunal section 21)
        if "discharge_risk" in self.popups:
            popup = self.popups["discharge_risk"]
            if hasattr(popup, 'set_notes_for_risk_analysis'):
                popup.set_notes_for_risk_analysis(raw_notes)
                print(f"[NURSING] Populated section 17 (discharge_risk) with {len(raw_notes)} notes for risk analysis")

        # Store pending data for popups not yet created
        # For community_support: prefer Social History, then Personal History, then raw notes
        community_entries = get_category_items("Social History")
        if not community_entries:
            community_entries = get_category_items("Personal History")
        if not community_entries:
            community_entries = raw_notes[:30] if raw_notes else []

        self._pending_section_data = {
            "community_support": community_entries,
            "progress": recent_progress,  # Use notes from last 6 months
            "awol": raw_notes,
            "risk_harm": raw_notes,  # TribunalRiskHarmPopup searches all notes
            "risk_property": raw_notes,  # TribunalRiskPropertyPopup searches all notes
            "seclusion": raw_notes,  # TribunalSeclusionPopup searches all notes
            "discharge_risk": raw_notes,  # GPRRiskPopup uses set_notes_for_risk_analysis
        }
        print(f"[NURSING] Stored pending data for sections not yet created")

        # Pre-fill medications in section 4 (nursing_care) - same as psych tribunal section 12
        if "nursing_care" in self.popups:
            self._prefill_medications_from_notes()

    def _prefill_medications_from_notes(self):
        """Extract medications from notes (last year only) and pre-fill section 4 with most recent per class."""
        import re
        from datetime import datetime, timedelta
        from CANONICAL_MEDS import MEDICATIONS

        if "nursing_care" not in self.popups:
            return

        popup = self.popups["nursing_care"]

        # Get ALL raw notes from page-level storage
        raw_notes = self._extracted_raw_notes
        if not raw_notes:
            print(f"[NURSING] No raw notes for medication extraction")
            return

        # === PARSE DATE HELPER ===
        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # === FIND MOST RECENT NOTE DATE AND FILTER TO LAST YEAR ===
        note_dates = [parse_date(n.get("date") or n.get("datetime")) for n in raw_notes]
        note_dates = [d for d in note_dates if d]
        if not note_dates:
            print(f"[NURSING] No dates found in notes")
            return

        latest_date = max(note_dates)
        cutoff_date = latest_date - timedelta(days=365)
        print(f"[NURSING] Latest note: {latest_date.date()}, searching from {cutoff_date.date()}")

        # Filter notes to last year only
        recent_notes = []
        for n in raw_notes:
            note_date = parse_date(n.get("date") or n.get("datetime"))
            if note_date and note_date >= cutoff_date:
                recent_notes.append(n)

        print(f"[NURSING] Extracting medications from {len(recent_notes)} notes (last year)...")

        # === BUILD TOKEN INDEX ===
        token_map = {}  # token -> (key, canonical)
        meta_map = {}   # key -> metadata

        for key, meta in MEDICATIONS.items():
            meta_map[key] = meta
            canonical = meta["canonical"]
            for syn in meta.get("patterns", []):
                s = syn.lower().strip()
                if s:
                    token_map[s] = (key, canonical)

        def tokenise(text):
            text = text.lower()
            text = re.sub(r'(?<=\d)\.(?=\d)', 'DOT', text)
            text = re.sub(r'[^a-z0-9DOT]+', ' ', text)
            text = text.replace("DOT", ".")
            return text.split()

        def find_dose(tokens, idx):
            unit_set = {"mg", "mcg", "µg", "g", "units", "iu"}
            for i in range(idx + 1, min(idx + 4, len(tokens))):
                tok = tokens[i]
                m = re.match(r'(\d+(?:\.\d+)?)(mg|mcg|µg|g|units|iu)$', tok)
                if m:
                    return float(m.group(1)), m.group(2)
                if tok.isdigit() or re.match(r'\d+\.\d+', tok):
                    if i + 1 < len(tokens) and tokens[i + 1] in unit_set:
                        return float(tok), tokens[i + 1]
            return None, None

        def find_freq(tokens, idx):
            freq_set = {"od", "bd", "tds", "qds", "qid", "nocte", "mane", "stat", "prn", "daily", "weekly", "monthly"}
            for i in range(idx, min(idx + 5, len(tokens))):
                if tokens[i] in freq_set:
                    return tokens[i]
            return None

        # === PSYCHIATRIC MEDICATION CLASSES (prioritized) ===
        PSYCH_CLASSES = {
            "Antipsychotic": 1,
            "Antidepressant": 2,
            "Mood Stabiliser": 3,
            "Anxiolytic": 4,
            "Benzodiazepine": 5,
            "Anticonvulsant / Benzodiazepine": 5,
            "Sedative": 6,
            "Sedation": 6,
            "Sleep": 7,
            "Stimulant": 8,
            "Addictions": 9,
            "Opioid Substitution": 9,
        }

        # === EXTRACT MEDICATIONS ===
        meds_found = []
        for n in recent_notes:
            content = n.get("content", "") or n.get("text", "") or ""
            if not content:
                continue
            tokens = tokenise(content)
            note_date = parse_date(n.get("date") or n.get("datetime"))

            for i, tok in enumerate(tokens):
                if tok in token_map:
                    key, canonical = token_map[tok]
                    meta = meta_map[key]
                    med_class = meta.get("class", "Other")

                    strength, unit = find_dose(tokens, i)
                    freq = find_freq(tokens, i)

                    meds_found.append({
                        "med_key": key,
                        "canonical": canonical,
                        "class": med_class,
                        "strength": strength,
                        "unit": unit or "mg",
                        "frequency": freq,
                        "date": note_date,
                    })

        print(f"[NURSING] Found {len(meds_found)} medication mentions in last year")

        if not meds_found:
            print(f"[NURSING] No medications found in notes")
            return

        # === GROUP BY CLASS AND PICK MOST RECENT PER CLASS ===
        class_groups = {}
        for med in meds_found:
            med_class = med.get("class", "Other")
            if med_class not in class_groups:
                class_groups[med_class] = []
            class_groups[med_class].append(med)

        # For each class, find the most recent medication with a dose
        most_recent_by_class = []
        for med_class, mentions in class_groups.items():
            # Sort by date descending
            sorted_mentions = sorted(
                mentions,
                key=lambda x: x.get("date") or datetime.min,
                reverse=True
            )
            # Take most recent with a dose
            for m in sorted_mentions:
                if m.get("strength"):
                    most_recent_by_class.append(m)
                    break
            else:
                # If none have dose, take most recent anyway
                if sorted_mentions:
                    most_recent_by_class.append(sorted_mentions[0])

        # === PRIORITIZE PSYCHIATRIC CLASSES ===
        def class_priority(med):
            med_class = med.get("class", "Other")
            return PSYCH_CLASSES.get(med_class, 100)  # Non-psych classes get low priority

        # Sort: psych classes first (by priority), then others
        most_recent_by_class.sort(key=class_priority)

        # Limit to top entries (psych meds first)
        max_meds = 10
        final_meds = most_recent_by_class[:max_meds]

        print(f"[NURSING] Selected {len(final_meds)} medications (1 per class, psych prioritized):")
        for m in final_meds:
            strength = m.get('strength')
            dose_str = f"{strength}{m.get('unit', 'mg')}" if strength else "no dose"
            print(f"[NURSING]   [{m.get('class')}] {m.get('canonical')}: {dose_str} {m.get('frequency') or ''} ({m.get('date').date() if m.get('date') else 'no date'})")

        # Map extracted frequency to popup frequency options
        FREQ_MAP = {
            "od": "OD", "daily": "OD", "once daily": "OD", "1x daily": "OD",
            "bd": "BD", "twice daily": "BD", "2x daily": "BD",
            "tds": "TDS", "three times daily": "TDS", "3x daily": "TDS",
            "qds": "QDS", "qid": "QDS", "four times daily": "QDS", "4x daily": "QDS",
            "nocte": "Nocte", "at night": "Nocte", "night": "Nocte", "on": "Nocte",
            "mane": "Mane", "in the morning": "Mane", "am": "Mane",
            "prn": "PRN", "as required": "PRN", "when required": "PRN",
            "weekly": "Weekly", "1 weekly": "Weekly", "once weekly": "Weekly",
            "fortnightly": "Fortnightly", "2 weekly": "Fortnightly", "every 2 weeks": "Fortnightly",
            "3 weekly": "3 Weekly", "every 3 weeks": "3 Weekly",
            "monthly": "Monthly", "4 weekly": "Monthly", "every 4 weeks": "Monthly",
        }

        # Populate the medication entries
        while len(popup._medications) < len(final_meds):
            popup._add_medication_entry()

        for i, med in enumerate(final_meds):
            if i >= len(popup._medications):
                break

            entry = popup._medications[i]
            med_key = med.get("med_key", "")

            # Set medication name (the key in MEDICATIONS dict)
            name_combo = entry.get("name")
            if name_combo:
                idx = name_combo.findText(med_key)
                if idx >= 0:
                    name_combo.setCurrentIndex(idx)
                else:
                    # Try canonical name
                    canonical = med.get("canonical", "")
                    idx = name_combo.findText(canonical.upper())
                    if idx >= 0:
                        name_combo.setCurrentIndex(idx)

            # Set dose
            dose_combo = entry.get("dose")
            if dose_combo and med.get("strength"):
                strength = med.get("strength")
                unit = med.get("unit", "mg")
                if isinstance(strength, float) and strength == int(strength):
                    dose_str = f"{int(strength)}{unit}"
                else:
                    dose_str = f"{strength}{unit}"
                idx = dose_combo.findText(dose_str)
                if idx >= 0:
                    dose_combo.setCurrentIndex(idx)
                else:
                    dose_combo.setCurrentText(dose_str)

            # Set frequency
            freq_combo = entry.get("freq")
            if freq_combo and med.get("frequency"):
                freq = med.get("frequency", "").lower()
                mapped_freq = FREQ_MAP.get(freq, "OD")
                idx = freq_combo.findText(mapped_freq)
                if idx >= 0:
                    freq_combo.setCurrentIndex(idx)

        print(f"[NURSING] ✓ Pre-filled {len(final_meds)} medication(s)")


# ================================================================
# NURSING-SPECIFIC POPUPS
# ================================================================

FREQUENCY_OPTIONS = ["OD", "BD", "TDS", "QDS", "Nocte", "PRN", "Weekly", "Fortnightly", "Monthly"]


class NursingCarePopup(QWidget):
    """Popup for nursing care and medication section with medication dropdowns."""

    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._possessive = "his" if gender == "Male" else "her"
        self._medications = []

        self._setup_ui()
        add_lock_to_popup(self, show_button=False)

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Scrollable content
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(12, 12, 12, 12)
        scroll_layout.setSpacing(12)
        scroll.setWidget(scroll_content)

        # === NURSING CARE SECTION ===
        care_label = QLabel("Nature of nursing care being provided:")
        care_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        scroll_layout.addWidget(care_label)

        self.care_dropdown = QComboBox()
        self.care_dropdown.addItems([
            "Select nursing care level...",
            "General inpatient nursing care",
            "Enhanced nursing observation and support",
            "1:1 nursing support",
            "2:1 nursing support",
            "Specialised psychiatric nursing care",
            "Rehabilitation nursing care"
        ])
        self.care_dropdown.setFixedHeight(40)
        self.care_dropdown.setStyleSheet("""
            QComboBox {
                background: #d1d5db;
                border: 1px solid #9ca3af;
                border-radius: 6px;
                padding: 8px 12px;
                font-size: 17px;
                min-height: 24px;
            }
            QComboBox:hover {
                border-color: #14b8a6;
            }
            QComboBox::drop-down {
                border: none;
                width: 24px;
            }
            QComboBox QAbstractItemView {
                background: white;
                border: 1px solid #9ca3af;
                selection-background-color: #14b8a6;
            }
        """)
        self.care_dropdown.currentIndexChanged.connect(self._send)
        scroll_layout.addWidget(self.care_dropdown)

        self.care_details = QTextEdit()
        self.care_details.setPlaceholderText("Additional nursing care details (optional)...")
        self.care_details.setMaximumHeight(60)
        self.care_details.setStyleSheet("""
            QTextEdit {
                background: #d1d5db;
                border: 1px solid #9ca3af;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
            }
            QTextEdit:focus {
                border-color: #14b8a6;
            }
        """)
        self.care_details.textChanged.connect(self._send)
        scroll_layout.addWidget(self.care_details)

        # === MEDICATION SECTION ===
        med_label = QLabel("Current Medication")
        med_label.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151; margin-top: 8px;")
        scroll_layout.addWidget(med_label)

        self.med_entries_container = QWidget()
        self.med_entries_layout = QVBoxLayout(self.med_entries_container)
        self.med_entries_layout.setContentsMargins(0, 0, 0, 0)
        self.med_entries_layout.setSpacing(4)
        scroll_layout.addWidget(self.med_entries_container)

        self._add_medication_entry()

        add_med_btn = QPushButton("+ Add Medication")
        add_med_btn.setStyleSheet("""
            QPushButton {
                background: #e5e7eb;
                color: #374151;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 17px;
            }
            QPushButton:hover {
                background: #d1d5db;
            }
        """)
        add_med_btn.clicked.connect(self._add_medication_entry)
        scroll_layout.addWidget(add_med_btn)

        scroll_layout.addStretch()
        layout.addWidget(scroll, 1)

    def _add_medication_entry(self):
        from CANONICAL_MEDS import MEDICATIONS

        # Minimal styling to preserve native dropdown arrow visibility
        combo_style = """
            QComboBox {
                background: white;
                border: 1px solid #6b7280;
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 17px;
                min-height: 20px;
            }
            QComboBox:hover {
                border-color: #14b8a6;
            }
        """

        entry_widget = QFrame()
        entry_widget.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border: 1px solid #9ca3af;
                border-radius: 6px;
            }
            QLabel {
                font-weight: 600;
                color: #374151;
            }
        """)
        entry_layout = QVBoxLayout(entry_widget)
        entry_layout.setContentsMargins(10, 8, 10, 8)
        entry_layout.setSpacing(6)

        # Medication name row
        name_row = QHBoxLayout()
        name_combo = QComboBox()
        name_combo.setEditable(True)
        name_combo.addItem("")
        name_combo.addItems(sorted(MEDICATIONS.keys()))
        name_combo.setMinimumWidth(180)
        name_combo.setFixedHeight(32)
        name_combo.setStyleSheet(combo_style)
        name_row.addWidget(QLabel("Med:"))
        name_row.addWidget(name_combo)
        name_row.addStretch()

        # Remove button
        remove_btn = QPushButton("X")
        remove_btn.setFixedSize(24, 24)
        remove_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 4px;
                font-size: 17px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: #dc2626;
            }
        """)
        name_row.addWidget(remove_btn)
        entry_layout.addLayout(name_row)

        # Dose row
        dose_row = QHBoxLayout()
        dose_combo = QComboBox()
        dose_combo.setEditable(True)
        dose_combo.setMinimumWidth(100)
        dose_combo.setFixedHeight(32)
        dose_combo.setStyleSheet(combo_style)
        dose_row.addWidget(QLabel("Dose:"))
        dose_row.addWidget(dose_combo)
        dose_row.addStretch()
        entry_layout.addLayout(dose_row)

        # Frequency row
        freq_row = QHBoxLayout()
        freq_combo = QComboBox()
        freq_combo.addItems(FREQUENCY_OPTIONS)
        freq_combo.setMinimumWidth(100)
        freq_combo.setFixedHeight(32)
        freq_combo.setStyleSheet(combo_style)
        freq_row.addWidget(QLabel("Freq:"))
        freq_row.addWidget(freq_combo)
        freq_row.addStretch()
        entry_layout.addLayout(freq_row)

        # BNF max label
        bnf_label = QLabel("")
        bnf_label.setStyleSheet("font-size: 16px; color: #666; font-style: italic;")
        entry_layout.addWidget(bnf_label)

        entry_data = {
            "widget": entry_widget,
            "name": name_combo,
            "dose": dose_combo,
            "freq": freq_combo,
            "bnf": bnf_label
        }
        self._medications.append(entry_data)

        def on_med_change(med_name):
            if med_name and med_name in MEDICATIONS:
                info = MEDICATIONS[med_name]
                allowed = info.get("allowed_strengths", [])
                dose_combo.clear()
                if allowed:
                    dose_combo.addItems([f"{s}mg" for s in allowed])
                bnf_max = info.get("bnf_max", "")
                bnf_label.setText(f"Max BNF: {bnf_max}" if bnf_max else "")
            else:
                dose_combo.clear()
                bnf_label.setText("")
            self._send()

        def remove_entry():
            if len(self._medications) > 1:
                self._medications.remove(entry_data)
                entry_widget.deleteLater()
                self._send()

        name_combo.currentTextChanged.connect(on_med_change)
        dose_combo.currentTextChanged.connect(self._send)
        freq_combo.currentIndexChanged.connect(self._send)
        remove_btn.clicked.connect(remove_entry)

        self.med_entries_layout.addWidget(entry_widget)

    def _generate_text(self) -> str:
        parts = []

        # Nursing care
        care_idx = self.care_dropdown.currentIndex()
        if care_idx > 0:
            care_text = self.care_dropdown.currentText()
            parts.append(f"Nursing care: {care_text}")

        care_details = self.care_details.toPlainText().strip()
        if care_details:
            if parts:
                parts[-1] += f". {care_details}"
            else:
                parts.append(f"Nursing care: {care_details}")

        # Medications
        med_lines = []
        for entry in self._medications:
            name = entry["name"].currentText().strip()
            dose = entry["dose"].currentText().strip()
            freq = entry["freq"].currentText().strip()
            if name:
                line = name.capitalize()
                if dose:
                    line += f" {dose}"
                if freq:
                    line += f" {freq}"
                med_lines.append(line)

        if med_lines:
            parts.append("Current medication: " + ", ".join(med_lines) + ".")

        return "\n\n".join(parts)

    def _send(self):
        text = self._generate_text()
        if text:
            self.sent.emit(text)


class ObservationLevelPopup(QWidget):
    """Popup for observation level section."""

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Level of observation patient is currently subject to:")
        label.setStyleSheet("font-size: 22px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        # Observation level options
        self.obs_combo = QComboBox()
        self.obs_combo.setFixedHeight(40)
        self.obs_combo.setStyleSheet("QComboBox { font-size: 22px; padding: 4px 8px; } QComboBox QAbstractItemView { font-size: 22px; }")
        self.obs_combo.addItems([
            "Select observation level...",
            "General observation",
            "Intermittent observation (15-30 minutes)",
            "Continuous observation (within eyesight)",
            "Within arm's length",
            "2:1 observation",
            "Other (specify below)"
        ])
        self.obs_combo.currentIndexChanged.connect(self._send)
        layout.addWidget(self.obs_combo)

        # Additional details
        self.details_edit = QTextEdit()
        self.details_edit.setStyleSheet("font-size: 22px;")
        self.details_edit.setPlaceholderText("Additional details about observation level...")
        self.details_edit.setMinimumHeight(80)
        self.details_edit.textChanged.connect(self._send)
        layout.addWidget(self.details_edit)

        layout.addStretch()

        add_lock_to_popup(self, show_button=False)

    def _send(self):
        level = self.obs_combo.currentText()
        details = self.details_edit.toPlainText().strip()

        if level and level != "Select observation level...":
            output = f"The patient is currently on {level.lower()}."
            if details:
                output += f" {details}"
            self.sent.emit(output)


CONTACT_LEVELS = ["low", "some", "moderate", "good", "significant"]
CONTACT_PHRASES = {
    "low": ["minimal contact with", "limited contact with", "occasional contact with"],
    "some": ["some contact with", "intermittent contact with", "periodic contact with"],
    "moderate": ["moderate contact with", "regular contact with", "reasonable contact with"],
    "good": ["good contact with", "consistent contact with", "frequent contact with"],
    "significant": ["significant contact with", "strong contact with", "close contact with"],
}

RELATIVE_TYPES = [
    "mother", "father", "stepmother", "stepfather",
    "brothers", "sisters", "aunt", "uncle", "cousin", "grandparents"
]
PLURAL_RELATIVES = ["brothers", "sisters", "aunt", "uncle", "cousin", "grandparents"]
COUNT_OPTIONS_LONG = ["1", "2", "3", "4", "5", "5+"]
COUNT_OPTIONS_SHORT = ["1", "2", "3", "4"]


class ContactPopup(QWidget):
    """Popup for contact with relatives/friends/other patients section."""

    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._possessive = "his" if gender == "Male" else "her"

        # State for relatives, friends, other patients
        self._relatives = []  # List of {"type": str, "count": str, "level": int}
        self._type_combos = []  # Track all type combo widgets for updating available options
        self._friends_level = None  # 0-4 index into CONTACT_LEVELS
        self._patients_level = None  # 0-4 index into CONTACT_LEVELS

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # ========== SCROLLABLE CONTENT ==========
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(16)

        # ========== SECTION 1: RELATIVES ==========
        relatives_section = QFrame()
        relatives_section.setStyleSheet("QFrame { background: #f0fdfa; border: 1px solid #99f6e4; border-radius: 8px; }")
        rel_layout = QVBoxLayout(relatives_section)
        rel_layout.setContentsMargins(12, 12, 12, 12)
        rel_layout.setSpacing(10)

        rel_header = QLabel("Relatives")
        rel_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        rel_layout.addWidget(rel_header)

        # Container for relative entries
        self._relatives_container = QVBoxLayout()
        self._relatives_container.setSpacing(8)
        rel_layout.addLayout(self._relatives_container)

        # Add initial empty entry
        self._add_relative_entry()

        # Add button
        add_rel_btn = QPushButton("+ Add Relative")
        add_rel_btn.setStyleSheet("""
            QPushButton { background: #14b8a6; color: white; border: none; padding: 8px 16px; border-radius: 6px; font-weight: 600; }
            QPushButton:hover { background: #0d9488; }
        """)
        add_rel_btn.clicked.connect(self._add_relative_entry)
        rel_layout.addWidget(add_rel_btn)

        scroll_layout.addWidget(relatives_section)

        # ========== SECTION 2: FRIENDS ==========
        friends_section = QFrame()
        friends_section.setStyleSheet("QFrame { background: #f0fdfa; border: 1px solid #99f6e4; border-radius: 8px; }")
        friends_layout = QVBoxLayout(friends_section)
        friends_layout.setContentsMargins(12, 12, 12, 12)
        friends_layout.setSpacing(10)

        friends_header = QLabel("Friends")
        friends_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        friends_layout.addWidget(friends_header)

        self._friends_slider_container = QWidget()
        self._friends_slider_container.setVisible(False)
        friends_slider_lay = QVBoxLayout(self._friends_slider_container)
        friends_slider_lay.setContentsMargins(0, 0, 0, 0)
        friends_slider_lay.setSpacing(4)

        friends_slider_label = QLabel("Contact level:")
        friends_slider_label.setStyleSheet("font-size: 17px; color: #374151;")
        friends_slider_lay.addWidget(friends_slider_label)

        self._friends_slider = NoWheelSlider(Qt.Horizontal)
        self._friends_slider.setRange(0, len(CONTACT_LEVELS) - 1)
        self._friends_slider.valueChanged.connect(self._on_friends_slider_change)
        self._friends_slider.setStyleSheet(self._slider_style())
        friends_slider_lay.addWidget(self._friends_slider)

        self._friends_level_label = QLabel("")
        self._friends_level_label.setStyleSheet("font-size: 16px; color: #0f766e; font-weight: 600;")
        friends_slider_lay.addWidget(self._friends_level_label)

        friends_layout.addWidget(self._friends_slider_container)

        # Has friends toggle
        friends_toggle_row = QHBoxLayout()
        self.friends_yes_btn = QPushButton("Has Friends")
        self.friends_no_btn = QPushButton("No Friends")
        for btn in [self.friends_yes_btn, self.friends_no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet(self._toggle_btn_style())
        self.friends_yes_btn.clicked.connect(self._on_friends_yes)
        self.friends_no_btn.clicked.connect(self._on_friends_no)
        friends_toggle_row.addWidget(self.friends_yes_btn)
        friends_toggle_row.addWidget(self.friends_no_btn)
        friends_toggle_row.addStretch()
        friends_layout.addLayout(friends_toggle_row)

        scroll_layout.addWidget(friends_section)

        # ========== SECTION 3: OTHER PATIENTS ==========
        patients_section = QFrame()
        patients_section.setStyleSheet("QFrame { background: #f0fdfa; border: 1px solid #99f6e4; border-radius: 8px; }")
        patients_layout = QVBoxLayout(patients_section)
        patients_layout.setContentsMargins(12, 12, 12, 12)
        patients_layout.setSpacing(10)

        patients_header = QLabel("Other Patients")
        patients_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        patients_layout.addWidget(patients_header)

        # Location checkboxes (ward/community)
        from PySide6.QtWidgets import QCheckBox
        location_row = QHBoxLayout()
        self.patients_ward_cb = QCheckBox("Ward")
        self.patients_community_cb = QCheckBox("Community")
        self.patients_ward_cb.setStyleSheet("QCheckBox { color: #374151; font-size: 22px; }")
        self.patients_community_cb.setStyleSheet("QCheckBox { color: #374151; font-size: 22px; }")
        self.patients_ward_cb.toggled.connect(self._on_patients_location_change)
        self.patients_community_cb.toggled.connect(self._on_patients_location_change)
        location_row.addWidget(self.patients_ward_cb)
        location_row.addSpacing(16)
        location_row.addWidget(self.patients_community_cb)
        location_row.addStretch()
        patients_layout.addLayout(location_row)

        self._patients_slider_container = QWidget()
        self._patients_slider_container.setVisible(False)
        patients_slider_lay = QVBoxLayout(self._patients_slider_container)
        patients_slider_lay.setContentsMargins(0, 0, 0, 0)
        patients_slider_lay.setSpacing(4)

        patients_slider_label = QLabel("Contact level:")
        patients_slider_label.setStyleSheet("font-size: 17px; color: #374151;")
        patients_slider_lay.addWidget(patients_slider_label)

        self._patients_slider = NoWheelSlider(Qt.Horizontal)
        self._patients_slider.setRange(0, len(CONTACT_LEVELS) - 1)
        self._patients_slider.valueChanged.connect(self._on_patients_slider_change)
        self._patients_slider.setStyleSheet(self._slider_style())
        patients_slider_lay.addWidget(self._patients_slider)

        self._patients_level_label = QLabel("")
        self._patients_level_label.setStyleSheet("font-size: 16px; color: #0f766e; font-weight: 600;")
        patients_slider_lay.addWidget(self._patients_level_label)

        patients_layout.addWidget(self._patients_slider_container)

        # Has patient contact toggle
        patients_toggle_row = QHBoxLayout()
        self.patients_yes_btn = QPushButton("Has Contact")
        self.patients_no_btn = QPushButton("No Contact")
        for btn in [self.patients_yes_btn, self.patients_no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet(self._toggle_btn_style())
        self.patients_yes_btn.clicked.connect(self._on_patients_yes)
        self.patients_no_btn.clicked.connect(self._on_patients_no)
        patients_toggle_row.addWidget(self.patients_yes_btn)
        patients_toggle_row.addWidget(self.patients_no_btn)
        patients_toggle_row.addStretch()
        patients_layout.addLayout(patients_toggle_row)

        scroll_layout.addWidget(patients_section)
        scroll_layout.addStretch()
        scroll.setWidget(scroll_content)
        layout.addWidget(scroll, 1)

    def _slider_style(self):
        return """
            QSlider::groove:horizontal {
                height: 8px;
                background: #d1d5db;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background: #14b8a6;
                width: 18px;
                height: 18px;
                margin: -5px 0;
                border-radius: 9px;
            }
            QSlider::sub-page:horizontal {
                background: #14b8a6;
                border-radius: 4px;
            }
        """

    def _toggle_btn_style(self):
        return """
            QPushButton {
                background: #f3f4f6;
                border: 1px solid #d1d5db;
                padding: 8px 16px;
                border-radius: 6px;
            }
            QPushButton:checked {
                background: #14b8a6;
                color: white;
                border-color: #14b8a6;
            }
        """

    def _add_relative_entry(self):
        import random
        entry_data = {"type": None, "count": None, "level": None, "widgets": {}}

        entry_widget = QFrame()
        entry_widget.setStyleSheet("QFrame { background: #e5e7eb; border: 1px solid #9ca3af; border-radius: 6px; }")
        entry_layout = QVBoxLayout(entry_widget)
        entry_layout.setContentsMargins(10, 8, 10, 8)
        entry_layout.setSpacing(6)

        # Top row: relative type dropdown and remove button
        top_row = QHBoxLayout()
        type_combo = QComboBox()
        type_combo.addItem("Select relative...")
        type_combo.addItems([r.capitalize() for r in RELATIVE_TYPES])
        type_combo.setMinimumWidth(150)
        type_combo.setFixedHeight(40)
        type_combo.setStyleSheet("""
            QComboBox { font-size: 22px; padding: 4px 8px; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        top_row.addWidget(type_combo)

        # Count dropdown (hidden by default)
        count_combo = QComboBox()
        count_combo.setMinimumWidth(60)
        count_combo.setFixedHeight(40)
        count_combo.setStyleSheet("""
            QComboBox { font-size: 22px; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        count_combo.setVisible(False)
        top_row.addWidget(count_combo)

        top_row.addStretch()

        remove_btn = QPushButton("X")
        remove_btn.setFixedSize(24, 24)
        remove_btn.setStyleSheet("""
            QPushButton { background: #ef4444; color: white; border: none; border-radius: 4px; font-size: 17px; font-weight: 600; }
            QPushButton:hover { background: #dc2626; }
        """)
        top_row.addWidget(remove_btn)
        entry_layout.addLayout(top_row)

        # Contact slider (hidden by default)
        slider_container = QWidget()
        slider_container.setVisible(False)
        slider_lay = QVBoxLayout(slider_container)
        slider_lay.setContentsMargins(0, 4, 0, 0)
        slider_lay.setSpacing(4)

        slider_label = QLabel("Contact level:")
        slider_label.setStyleSheet("font-size: 17px; color: #374151;")
        slider_lay.addWidget(slider_label)

        contact_slider = NoWheelSlider(Qt.Horizontal)
        contact_slider.setRange(0, len(CONTACT_LEVELS) - 1)
        contact_slider.setStyleSheet(self._slider_style())
        slider_lay.addWidget(contact_slider)

        level_label = QLabel("")
        level_label.setStyleSheet("font-size: 16px; color: #0f766e; font-weight: 600;")
        slider_lay.addWidget(level_label)

        entry_layout.addWidget(slider_container)

        entry_data["widgets"] = {
            "frame": entry_widget,
            "type_combo": type_combo,
            "count_combo": count_combo,
            "slider_container": slider_container,
            "contact_slider": contact_slider,
            "level_label": level_label,
        }

        def on_type_change(idx):
            if idx == 0:
                entry_data["type"] = None
                entry_data["count"] = None
                entry_data["level"] = None
                count_combo.setVisible(False)
                slider_container.setVisible(False)
            else:
                # Get the selected text and convert to lowercase for the type
                selected_text = type_combo.currentText()
                rel_type = selected_text.lower()
                entry_data["type"] = rel_type

                # Show count dropdown for plural relatives
                if rel_type in PLURAL_RELATIVES:
                    count_combo.clear()
                    if rel_type in ["brothers", "sisters"]:
                        count_combo.addItems(COUNT_OPTIONS_LONG)
                    else:
                        count_combo.addItems(COUNT_OPTIONS_SHORT)
                    count_combo.setVisible(True)
                    entry_data["count"] = count_combo.currentText()
                else:
                    count_combo.setVisible(False)
                    entry_data["count"] = None

                # Show slider
                slider_container.setVisible(True)
                entry_data["level"] = contact_slider.value()
                level_label.setText(CONTACT_LEVELS[contact_slider.value()].capitalize())

            self._update_available_relatives()
            self._send()

        def on_count_change(idx):
            if count_combo.isVisible():
                entry_data["count"] = count_combo.currentText()
            self._send()

        def on_slider_change(val):
            entry_data["level"] = val
            level_label.setText(CONTACT_LEVELS[val].capitalize())
            self._send()

        def on_remove():
            if entry_data in self._relatives:
                self._relatives.remove(entry_data)
            if type_combo in self._type_combos:
                self._type_combos.remove(type_combo)
            entry_widget.deleteLater()
            self._update_available_relatives()
            self._send()

        type_combo.currentIndexChanged.connect(on_type_change)
        count_combo.currentIndexChanged.connect(on_count_change)
        contact_slider.valueChanged.connect(on_slider_change)
        remove_btn.clicked.connect(on_remove)

        self._relatives.append(entry_data)
        self._type_combos.append(type_combo)
        self._relatives_container.addWidget(entry_widget)
        self._update_available_relatives()

    def _update_available_relatives(self):
        """Update all type combos to hide already-selected relatives."""
        # Get all currently selected types
        selected_types = set()
        for entry in self._relatives:
            if entry.get("type"):
                selected_types.add(entry["type"])

        # Update each combo to show only available options
        for combo in self._type_combos:
            current_type = None
            current_idx = combo.currentIndex()
            if current_idx > 0:
                # Get the actual type from the entry data
                for entry in self._relatives:
                    if entry.get("widgets", {}).get("type_combo") == combo:
                        current_type = entry.get("type")
                        break

            # Block signals while updating
            combo.blockSignals(True)
            combo.clear()
            combo.addItem("Select relative...")

            for rel_type in RELATIVE_TYPES:
                # Show if: not selected by others, OR is our own selection
                if rel_type not in selected_types or rel_type == current_type:
                    combo.addItem(rel_type.capitalize())

            # Restore selection if we had one
            if current_type:
                for i in range(combo.count()):
                    if combo.itemText(i).lower() == current_type:
                        combo.setCurrentIndex(i)
                        break

            combo.blockSignals(False)

    def _on_friends_yes(self):
        self.friends_no_btn.setChecked(False)
        self.friends_yes_btn.setChecked(True)
        self._friends_slider_container.setVisible(True)
        self._friends_level = self._friends_slider.value()
        self._friends_level_label.setText(CONTACT_LEVELS[self._friends_level].capitalize())
        self._send()

    def _on_friends_no(self):
        self.friends_yes_btn.setChecked(False)
        self.friends_no_btn.setChecked(True)
        self._friends_slider_container.setVisible(False)
        self._friends_level = -1  # Indicates "no friends"
        self._send()

    def _on_friends_slider_change(self, val):
        self._friends_level = val
        self._friends_level_label.setText(CONTACT_LEVELS[val].capitalize())
        self._send()

    def _on_patients_yes(self):
        self.patients_no_btn.setChecked(False)
        self.patients_yes_btn.setChecked(True)
        self._patients_slider_container.setVisible(True)
        self._patients_level = self._patients_slider.value()
        self._patients_level_label.setText(CONTACT_LEVELS[self._patients_level].capitalize())
        self._send()

    def _on_patients_no(self):
        self.patients_yes_btn.setChecked(False)
        self.patients_no_btn.setChecked(True)
        self._patients_slider_container.setVisible(False)
        self._patients_level = -1  # Indicates "no contact"
        self._send()

    def _on_patients_slider_change(self, val):
        self._patients_level = val
        self._patients_level_label.setText(CONTACT_LEVELS[val].capitalize())
        self._send()

    def _on_patients_location_change(self, checked):
        self._send()

    def _format_relative(self, entry):
        import random
        rel_type = entry.get("type")
        count = entry.get("count")
        level = entry.get("level")

        if not rel_type or level is None:
            return None

        level_key = CONTACT_LEVELS[level]
        phrase = random.choice(CONTACT_PHRASES[level_key])

        # Format the relative description
        if rel_type in ["mother", "father", "stepmother", "stepfather"]:
            return f"{phrase} {self._possessive} {rel_type}"
        elif rel_type == "brothers":
            if count == "1":
                return f"{phrase} {self._possessive} brother"
            else:
                return f"{phrase} {self._possessive} {count} brothers"
        elif rel_type == "sisters":
            if count == "1":
                return f"{phrase} {self._possessive} sister"
            else:
                return f"{phrase} {self._possessive} {count} sisters"
        elif rel_type == "aunt":
            if count == "1":
                return f"{phrase} {self._possessive} aunt"
            else:
                return f"{phrase} {self._possessive} {count} aunts"
        elif rel_type == "uncle":
            if count == "1":
                return f"{phrase} {self._possessive} uncle"
            else:
                return f"{phrase} {self._possessive} {count} uncles"
        elif rel_type == "cousin":
            if count == "1":
                return f"{phrase} {self._possessive} cousin"
            else:
                return f"{phrase} {self._possessive} {count} cousins"
        elif rel_type == "grandparents":
            if count == "1":
                return f"{phrase} {self._possessive} grandparent"
            else:
                return f"{phrase} {self._possessive} {count} grandparents"
        return None

    def _formatted_text(self):
        import random
        parts = []

        # Relatives
        rel_parts = []
        for entry in self._relatives:
            formatted = self._format_relative(entry)
            if formatted:
                rel_parts.append(formatted)

        if rel_parts:
            if len(rel_parts) == 1:
                parts.append(f"{self._pronoun} has {rel_parts[0]}.")
            else:
                joined = ", ".join(rel_parts[:-1]) + " and " + rel_parts[-1]
                parts.append(f"{self._pronoun} has {joined}.")

        # Friends
        if self._friends_level is not None:
            if self._friends_level == -1:
                parts.append(f"{self._pronoun} does not have contact with friends.")
            else:
                level_key = CONTACT_LEVELS[self._friends_level]
                phrase = random.choice(CONTACT_PHRASES[level_key])
                parts.append(f"{self._pronoun} has {phrase} friends.")

        # Other patients - check ward/community location
        is_ward = self.patients_ward_cb.isChecked()
        is_community = self.patients_community_cb.isChecked()

        if self._patients_level is not None:
            if self._patients_level == -1:
                if is_ward and is_community:
                    parts.append(f"{self._pronoun} does not interact with other patients on the ward or in the community.")
                elif is_ward:
                    parts.append(f"{self._pronoun} does not interact with other patients on the ward.")
                elif is_community:
                    parts.append(f"{self._pronoun} does not interact with other patients in the community.")
                else:
                    parts.append(f"{self._pronoun} does not interact with other patients.")
            else:
                level_key = CONTACT_LEVELS[self._patients_level]
                phrase = random.choice(CONTACT_PHRASES[level_key])
                if is_ward and is_community:
                    parts.append(f"{self._pronoun} has {phrase} other patients both on the ward and in the community.")
                elif is_ward:
                    parts.append(f"{self._pronoun} has {phrase} other patients on the ward.")
                elif is_community:
                    parts.append(f"{self._pronoun} has {phrase} other patients in the community.")
                else:
                    parts.append(f"{self._pronoun} has {phrase} other patients.")

        return " ".join(parts)

    def set_gender(self, gender: str):
        """Update gender and pronouns."""
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._possessive = "his" if gender == "Male" else "her"
        self._send()
        add_lock_to_popup(self, show_button=False)

    def _send(self):
        text = self._formatted_text()
        if text:
            self.sent.emit(text)


SUPPORT_LEVELS = ["minimal", "some", "moderate", "good", "significant"]
SUPPORT_PHRASES = {
    "minimal": ["minimally supported by", "receives limited support from", "has occasional support from"],
    "some": ["receives some support from", "is somewhat supported by", "has intermittent support from"],
    "moderate": ["receives moderate support from", "is reasonably supported by", "has regular support from"],
    "good": ["is well supported by", "receives good support from", "has consistent support from"],
    "significant": ["is mainly supported by", "receives significant support from", "has strong support from"],
}

FLOATING_SUPPORT_OPTIONS = [
    "Select frequency...",
    "24 hour",
    "4x/day",
    "3x/day",
    "2x/day",
    "Daily",
    "Every other day",
    "Every 2 days",
    "Twice a week",
    "Once a week",
    "Every 2 weeks"
]


class CommunitySupportPopup(QWidget):
    """Popup for community support section with family support, community team, and accommodation."""

    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._possessive = "his" if gender == "Male" else "her"
        self._obj = "him" if gender == "Male" else "her"

        # State
        self._supporters = []  # List of {"type": str, "count": str, "level": int}

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # ========== SCROLLABLE CONTENT ==========
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(16)

        # ========== PANEL 1: FAMILY SUPPORT, COMMUNITY TEAM, ACCOMMODATION ==========
        panel1 = QFrame()
        panel1.setStyleSheet("QFrame { background: #f0fdfa; border: 1px solid #99f6e4; border-radius: 8px; }")
        panel1_layout = QVBoxLayout(panel1)
        panel1_layout.setContentsMargins(12, 12, 12, 12)
        panel1_layout.setSpacing(12)

        # --- FAMILY SUPPORT SECTION ---
        family_header = QLabel("Family Support in Community")
        family_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        panel1_layout.addWidget(family_header)

        self._supporters_container = QVBoxLayout()
        self._supporters_container.setSpacing(8)
        panel1_layout.addLayout(self._supporters_container)

        self._add_supporter_entry()

        add_support_btn = QPushButton("+ Add Supporter")
        add_support_btn.setStyleSheet("""
            QPushButton { background: #14b8a6; color: white; border: none; padding: 8px 16px; border-radius: 6px; font-weight: 600; }
            QPushButton:hover { background: #0d9488; }
        """)
        add_support_btn.clicked.connect(self._add_supporter_entry)
        panel1_layout.addWidget(add_support_btn)

        # Divider
        divider1 = QFrame()
        divider1.setFrameShape(QFrame.HLine)
        divider1.setStyleSheet("background: #99f6e4;")
        divider1.setFixedHeight(1)
        panel1_layout.addWidget(divider1)

        # --- COMMUNITY TEAM SECTION ---
        team_header = QLabel("Community Team")
        team_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        panel1_layout.addWidget(team_header)

        from PySide6.QtWidgets import QCheckBox
        self.cmht_cb = QCheckBox("CMHT (Community Mental Health Team)")
        self.cmht_cb.toggled.connect(self._send)
        self.cmht_cb.setStyleSheet("QCheckBox { color: #374151; font-size: 22px; }")
        panel1_layout.addWidget(self.cmht_cb)

        self.treatment_plan_cb = QCheckBox("Treatment Plan")
        self.treatment_plan_cb.toggled.connect(self._send)
        self.treatment_plan_cb.setStyleSheet("QCheckBox { color: #374151; font-size: 22px; }")
        panel1_layout.addWidget(self.treatment_plan_cb)

        # Divider
        divider2 = QFrame()
        divider2.setFrameShape(QFrame.HLine)
        divider2.setStyleSheet("background: #99f6e4;")
        divider2.setFixedHeight(1)
        panel1_layout.addWidget(divider2)

        # --- ACCOMMODATION SECTION ---
        accom_header = QLabel("Accommodation")
        accom_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        panel1_layout.addWidget(accom_header)

        from PySide6.QtWidgets import QRadioButton, QButtonGroup
        self.accom_btn_group = QButtonGroup(self)

        self.accom_24hr_rb = QRadioButton("24 hour supported")
        self.accom_24hr_rb.setStyleSheet("""
            QRadioButton {
                font-size: 22px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.accom_24hr_rb.toggled.connect(self._on_accom_toggled)
        self.accom_btn_group.addButton(self.accom_24hr_rb, 0)
        panel1_layout.addWidget(self.accom_24hr_rb)

        self.accom_9to5_rb = QRadioButton("9-5 supported")
        self.accom_9to5_rb.setStyleSheet("""
            QRadioButton {
                font-size: 22px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.accom_9to5_rb.toggled.connect(self._on_accom_toggled)
        self.accom_btn_group.addButton(self.accom_9to5_rb, 1)
        panel1_layout.addWidget(self.accom_9to5_rb)

        self.accom_independent_rb = QRadioButton("Independent")
        self.accom_independent_rb.setStyleSheet("""
            QRadioButton {
                font-size: 22px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.accom_independent_rb.toggled.connect(self._on_accom_toggled)
        self.accom_btn_group.addButton(self.accom_independent_rb, 2)
        panel1_layout.addWidget(self.accom_independent_rb)

        # Floating support container
        self.floating_container = QWidget()
        floating_layout = QVBoxLayout(self.floating_container)
        floating_layout.setContentsMargins(20, 4, 0, 0)
        floating_layout.setSpacing(4)

        self.floating_cb = QCheckBox("Floating support")
        self.floating_cb.toggled.connect(self._on_floating_toggled)
        floating_layout.addWidget(self.floating_cb)

        self.floating_dropdown_container = QWidget()
        fd_layout = QHBoxLayout(self.floating_dropdown_container)
        fd_layout.setContentsMargins(20, 0, 0, 0)

        self.floating_dropdown = QComboBox()
        self.floating_dropdown.addItems(FLOATING_SUPPORT_OPTIONS)
        self.floating_dropdown.currentIndexChanged.connect(self._send)
        fd_layout.addWidget(self.floating_dropdown)
        fd_layout.addStretch()

        self.floating_dropdown_container.hide()
        floating_layout.addWidget(self.floating_dropdown_container)
        self.floating_container.hide()
        panel1_layout.addWidget(self.floating_container)

        self.accom_family_rb = QRadioButton("Family")
        self.accom_family_rb.setStyleSheet("""
            QRadioButton {
                font-size: 22px;
                background: transparent;
            }
            QRadioButton::indicator {
                width: 18px;
                height: 18px;
            }
        """)
        self.accom_family_rb.toggled.connect(self._on_accom_toggled)
        self.accom_btn_group.addButton(self.accom_family_rb, 3)
        panel1_layout.addWidget(self.accom_family_rb)

        scroll_layout.addWidget(panel1)

        # ========== PANEL 2: DATA EXTRACTOR ==========
        panel2 = QFrame()
        panel2.setStyleSheet("QFrame { background: #ecfdf5; border: 1px solid #a7f3d0; border-radius: 8px; }")
        panel2_layout = QVBoxLayout(panel2)
        panel2_layout.setContentsMargins(12, 12, 12, 12)
        panel2_layout.setSpacing(8)

        data_header = QLabel("Additional Information from Import")
        data_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        panel2_layout.addWidget(data_header)

        data_hint = QLabel("Use Import File from the toolbar to extract and categorize data from clinical notes. Relevant data will appear here.")
        data_hint.setWordWrap(True)
        data_hint.setStyleSheet("font-size: 16px; color: #6b7280; font-style: italic;")
        panel2_layout.addWidget(data_hint)

        self.data_extractor_text = QTextEdit()
        self.data_extractor_text.setPlaceholderText("Paste or type additional community support information here...")
        self.data_extractor_text.setMinimumHeight(100)
        self.data_extractor_text.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
            }
        """)
        self.data_extractor_text.textChanged.connect(self._send)
        panel2_layout.addWidget(self.data_extractor_text)

        scroll_layout.addWidget(panel2)
        scroll_layout.addStretch()
        scroll.setWidget(scroll_content)
        layout.addWidget(scroll, 1)

        add_lock_to_popup(self, show_button=False)

    def _slider_style(self):
        return """
            QSlider::groove:horizontal {
                height: 8px;
                background: #d1d5db;
                border-radius: 4px;
            }
            QSlider::handle:horizontal {
                background: #14b8a6;
                width: 18px;
                height: 18px;
                margin: -5px 0;
                border-radius: 9px;
            }
            QSlider::sub-page:horizontal {
                background: #14b8a6;
                border-radius: 4px;
            }
        """

    def _add_supporter_entry(self):
        entry_data = {"type": None, "count": None, "level": None, "widgets": {}}

        entry_widget = QFrame()
        entry_widget.setStyleSheet("QFrame { background: #e5e7eb; border: 1px solid #9ca3af; border-radius: 6px; }")
        entry_layout = QVBoxLayout(entry_widget)
        entry_layout.setContentsMargins(10, 8, 10, 8)
        entry_layout.setSpacing(6)

        top_row = QHBoxLayout()
        type_combo = QComboBox()
        type_combo.addItem("Select supporter...")
        type_combo.addItems([r.capitalize() for r in RELATIVE_TYPES])
        type_combo.setMinimumWidth(150)
        type_combo.setFixedHeight(40)
        type_combo.setStyleSheet("""
            QComboBox { font-size: 22px; padding: 4px 8px; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        top_row.addWidget(type_combo)

        count_combo = QComboBox()
        count_combo.setMinimumWidth(60)
        count_combo.setFixedHeight(40)
        count_combo.setStyleSheet("""
            QComboBox { font-size: 22px; }
            QComboBox QAbstractItemView { font-size: 22px; }
        """)
        count_combo.setVisible(False)
        top_row.addWidget(count_combo)

        top_row.addStretch()

        remove_btn = QPushButton("X")
        remove_btn.setFixedSize(24, 24)
        remove_btn.setStyleSheet("""
            QPushButton { background: #ef4444; color: white; border: none; border-radius: 4px; font-size: 17px; font-weight: 600; }
            QPushButton:hover { background: #dc2626; }
        """)
        top_row.addWidget(remove_btn)
        entry_layout.addLayout(top_row)

        slider_container = QWidget()
        slider_container.setVisible(False)
        slider_lay = QVBoxLayout(slider_container)
        slider_lay.setContentsMargins(0, 4, 0, 0)
        slider_lay.setSpacing(4)

        slider_label = QLabel("Support level:")
        slider_label.setStyleSheet("font-size: 17px; color: #374151;")
        slider_lay.addWidget(slider_label)

        support_slider = NoWheelSlider(Qt.Horizontal)
        support_slider.setRange(0, len(SUPPORT_LEVELS) - 1)
        support_slider.setStyleSheet(self._slider_style())
        slider_lay.addWidget(support_slider)

        level_label = QLabel("")
        level_label.setStyleSheet("font-size: 16px; color: #0f766e; font-weight: 600;")
        slider_lay.addWidget(level_label)

        entry_layout.addWidget(slider_container)

        entry_data["widgets"] = {
            "frame": entry_widget,
            "type_combo": type_combo,
            "count_combo": count_combo,
            "slider_container": slider_container,
            "support_slider": support_slider,
            "level_label": level_label,
        }

        def on_type_change(idx):
            if idx == 0:
                entry_data["type"] = None
                entry_data["count"] = None
                entry_data["level"] = None
                count_combo.setVisible(False)
                slider_container.setVisible(False)
            else:
                rel_type = RELATIVE_TYPES[idx - 1]
                entry_data["type"] = rel_type

                if rel_type in PLURAL_RELATIVES:
                    count_combo.clear()
                    if rel_type in ["brothers", "sisters"]:
                        count_combo.addItems(COUNT_OPTIONS_LONG)
                    else:
                        count_combo.addItems(COUNT_OPTIONS_SHORT)
                    count_combo.setVisible(True)
                    entry_data["count"] = count_combo.currentText()
                else:
                    count_combo.setVisible(False)
                    entry_data["count"] = None

                slider_container.setVisible(True)
                entry_data["level"] = support_slider.value()
                level_label.setText(SUPPORT_LEVELS[support_slider.value()].capitalize())

            self._send()

        def on_count_change(idx):
            if count_combo.isVisible():
                entry_data["count"] = count_combo.currentText()
            self._send()

        def on_slider_change(val):
            entry_data["level"] = val
            level_label.setText(SUPPORT_LEVELS[val].capitalize())
            self._send()

        def on_remove():
            if entry_data in self._supporters:
                self._supporters.remove(entry_data)
            entry_widget.deleteLater()
            self._send()

        type_combo.currentIndexChanged.connect(on_type_change)
        count_combo.currentIndexChanged.connect(on_count_change)
        support_slider.valueChanged.connect(on_slider_change)
        remove_btn.clicked.connect(on_remove)

        self._supporters.append(entry_data)
        self._supporters_container.addWidget(entry_widget)

    def _on_accom_toggled(self, checked):
        self.floating_container.setVisible(self.accom_independent_rb.isChecked())
        if not self.accom_independent_rb.isChecked():
            self.floating_cb.setChecked(False)
        self._send()

    def _on_floating_toggled(self, checked):
        self.floating_dropdown_container.setVisible(checked)
        if not checked:
            self.floating_dropdown.setCurrentIndex(0)
        self._send()

    def _format_supporter(self, entry):
        import random
        rel_type = entry.get("type")
        count = entry.get("count")
        level = entry.get("level")

        if not rel_type or level is None:
            return None, None

        level_key = SUPPORT_LEVELS[level]

        # Format the relative description
        if rel_type in ["mother", "father", "stepmother", "stepfather"]:
            relative_str = f"{self._possessive} {rel_type}"
        elif rel_type == "brothers":
            relative_str = f"{self._possessive} brother" if count == "1" else f"{self._possessive} {count} brothers"
        elif rel_type == "sisters":
            relative_str = f"{self._possessive} sister" if count == "1" else f"{self._possessive} {count} sisters"
        elif rel_type == "aunt":
            relative_str = f"{self._possessive} aunt" if count == "1" else f"{self._possessive} {count} aunts"
        elif rel_type == "uncle":
            relative_str = f"{self._possessive} uncle" if count == "1" else f"{self._possessive} {count} uncles"
        elif rel_type == "cousin":
            relative_str = f"{self._possessive} cousin" if count == "1" else f"{self._possessive} {count} cousins"
        elif rel_type == "grandparents":
            relative_str = f"{self._possessive} grandparent" if count == "1" else f"{self._possessive} {count} grandparents"
        else:
            return None, None

        return relative_str, level

    def _formatted_text(self):
        import random
        parts = []

        # Format family support with natural language (avoiding repetition)
        support_entries = []
        for entry in self._supporters:
            formatted = self._format_supporter(entry)
            if formatted[0]:
                support_entries.append(formatted)

        if support_entries:
            # Sort by level (highest first) for natural sentence construction
            support_entries.sort(key=lambda x: x[1], reverse=True)

            # Build natural sentence avoiding repetition
            if len(support_entries) == 1:
                rel_str, level = support_entries[0]
                level_key = SUPPORT_LEVELS[level]
                phrase = random.choice(SUPPORT_PHRASES[level_key])
                parts.append(f"In the community, {self._pronoun.lower()} {phrase} {rel_str}.")
            else:
                # Group by support level for more natural output
                sentence_parts = []
                for i, (rel_str, level) in enumerate(support_entries):
                    level_key = SUPPORT_LEVELS[level]
                    if i == 0:
                        # First entry - use full phrase
                        phrase = random.choice(SUPPORT_PHRASES[level_key])
                        sentence_parts.append(f"{phrase} {rel_str}")
                    else:
                        # Subsequent entries - use shorter connector based on level
                        if level >= 3:  # good or significant
                            sentence_parts.append(f"also well by {rel_str}")
                        elif level >= 2:  # moderate
                            sentence_parts.append(f"moderately by {rel_str}")
                        elif level >= 1:  # some
                            sentence_parts.append(f"somewhat by {rel_str}")
                        else:  # minimal
                            sentence_parts.append(f"minimally by {rel_str}")

                # Join with natural connectors
                if len(sentence_parts) == 2:
                    joined = f"{sentence_parts[0]} and {sentence_parts[1]}"
                else:
                    joined = ", ".join(sentence_parts[:-1]) + f", and {sentence_parts[-1]}"
                parts.append(f"In the community, {self._pronoun.lower()} is {joined}.")

        # Community team
        if self.cmht_cb.isChecked():
            parts.append("Involvement with a community team would be essential for effective management on discharge.")

        if self.treatment_plan_cb.isChecked():
            parts.append(f"{self._possessive.capitalize()} treatment plan in the community would include maintenance of medication, engagement with community psychological work, occupational therapy input and care-coordinator input following the CPA process.")

        # Accommodation
        if self.accom_24hr_rb.isChecked():
            parts.append(f"With respect to community residence, {self._pronoun.lower()} would require 24 hour supported accommodation with input from staff to monitor {self._possessive} mental state and compliance.")
        elif self.accom_9to5_rb.isChecked():
            parts.append(f"With respect to community residence, {self._pronoun.lower()} would require 9-5 supported accommodation with input from staff to monitor {self._possessive} mental state and compliance.")
        elif self.accom_independent_rb.isChecked():
            parts.append(f"With respect to community residence, {self._pronoun.lower()} is able to move into independent accommodation.")
            if self.floating_cb.isChecked():
                freq = self.floating_dropdown.currentText()
                if freq and freq != "Select frequency...":
                    parts.append(f"{self._pronoun} would require floating support {freq.lower()}.")
        elif self.accom_family_rb.isChecked():
            parts.append(f"With respect to community residence, {self._pronoun.lower()} would return to live with family.")

        # Data extractor content
        extra_text = self.data_extractor_text.toPlainText().strip()
        if extra_text:
            parts.append(extra_text)

        return " ".join(parts)

    def set_gender(self, gender: str):
        """Update gender and pronouns."""
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._possessive = "his" if gender == "Male" else "her"
        self._obj = "him" if gender == "Male" else "her"
        self._send()

    def _send(self):
        text = self._formatted_text()
        if text:
            self.sent.emit(text)

    def set_extracted_data(self, text):
        """Set text from data extractor."""
        self.data_extractor_text.setPlainText(text)
        self._send()

    def set_entries(self, entries: list, date_range_info: str = ""):
        """Set entries from data extractor - format and display in text area."""
        if not entries:
            return

        # Format entries as text for the data extractor panel
        content_parts = []
        for entry in entries[:20]:  # Limit to 20 entries
            content = entry.get('content', '') or entry.get('text', '')
            date = entry.get('date', '') or entry.get('datetime', '')
            if content:
                if date:
                    content_parts.append(f"[{str(date)[:10]}] {content[:300]}")
                else:
                    content_parts.append(content[:300])

        if content_parts:
            self.data_extractor_text.setPlainText('\n\n'.join(content_parts))
            self._send()
            print(f"[NURSING] CommunitySupportPopup received {len(entries)} entries")


AWOL_SEARCH_TERMS = [
    "awol",
    "absent without leave",
    "escaped",
    "went missing",
    "gone missing",
    "missing from ward",
    "missing from the ward",
    "failed to return",
    "did not return",
    "absconded",
    "abscond",
]

AWOL_EXCLUDE_TERMS = [
    "no risk of awol",
    "no awol",
    "nil awol",
    "not awol",
    "risk of awol low",
    "low risk of awol",
    "no history of awol",
    "no previous awol",
    "never been awol",
    "has not been awol",
    "has not gone awol",
    "no episodes of awol",
    "no incidents of awol",
    "denies awol",
    "no absconding",
    "nil absconding",
    "no escape",
    "nil escape",
    "no missing",
    "not missing",
    "returned safely",
    "returned on time",
]

# Prefixes to exclude from AWOL search (risk assessment lines)
AWOL_EXCLUDE_PREFIXES = [
    "risk:",
    "risk -",
    "risk-",
    "risk to self:",
    "risk to self -",
    "risk to self-",
    "risk to safety:",
    "risk to safety -",
    "risk to safety-",
    "risk to others:",
    "risk to others -",
    "risk to others-",
]

# Words that negate AWOL if they appear before the keyword
AWOL_NEGATION_WORDS = ["no", "nothing", "didn't", "did not", "never", "nil", "not", "denies", "denied"]


class AWOLPopup(QWidget):
    """Popup for AWOL/failed return section with search functionality."""

    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._possessive = "his" if gender == "Male" else "her"
        self._awol_results = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # ========== SCROLLABLE CONTENT ==========
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 0, 0)
        scroll_layout.setSpacing(16)

        # ========== PANEL 1: MANUAL ENTRY ==========
        panel1 = QFrame()
        panel1.setStyleSheet("QFrame { background: #f0fdfa; border: 1px solid #99f6e4; border-radius: 8px; }")
        panel1_layout = QVBoxLayout(panel1)
        panel1_layout.setContentsMargins(12, 12, 12, 12)
        panel1_layout.setSpacing(10)

        panel1_header = QLabel("AWOL Status")
        panel1_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #0f766e;")
        panel1_layout.addWidget(panel1_header)

        label = QLabel("Has the patient been absent without leave or failed to return from granted leave?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 17px; color: #374151;")
        panel1_layout.addWidget(label)

        # Yes/No
        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes - AWOL incidents")
        self.no_btn = QPushButton("No AWOL incidents")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6;
                    border: 1px solid #d1d5db;
                    padding: 8px 16px;
                    border-radius: 6px;
                }
                QPushButton:checked {
                    background: #14b8a6;
                    color: white;
                    border-color: #14b8a6;
                }
            """)
        self.yes_btn.clicked.connect(self._on_yes)
        self.no_btn.clicked.connect(self._on_no)
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        panel1_layout.addLayout(btn_row)

        # Details
        self.details_edit = QTextEdit()
        self.details_edit.setPlaceholderText("Provide details of AWOL concerns...")
        self.details_edit.setMinimumHeight(80)
        self.details_edit.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px;
                font-size: 17px;
            }
        """)
        self.details_edit.textChanged.connect(self._send)
        panel1_layout.addWidget(self.details_edit)

        scroll_layout.addWidget(panel1)

        # ========== PANEL 2: AWOL SEARCH RESULTS ==========
        panel2 = QFrame()
        panel2.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #fcd34d; border-radius: 8px; }")
        panel2_layout = QVBoxLayout(panel2)
        panel2_layout.setContentsMargins(12, 12, 12, 12)
        panel2_layout.setSpacing(8)

        panel2_header = QLabel("AWOL Episodes from Notes")
        panel2_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #92400e;")
        panel2_layout.addWidget(panel2_header)

        search_hint = QLabel("Searching notes for: AWOL, escaped, went missing, absconded, failed to return...")
        search_hint.setWordWrap(True)
        search_hint.setStyleSheet("font-size: 16px; color: #78716c; font-style: italic;")
        panel2_layout.addWidget(search_hint)

        self.awol_count_label = QLabel("No notes loaded yet")
        self.awol_count_label.setStyleSheet("font-size: 17px; color: #92400e; font-weight: 600;")
        panel2_layout.addWidget(self.awol_count_label)

        # Select All / Deselect All buttons
        select_row = QHBoxLayout()
        self.select_all_btn = QPushButton("Select All")
        self.select_all_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        self.select_all_btn.clicked.connect(self._select_all)
        select_row.addWidget(self.select_all_btn)

        self.deselect_all_btn = QPushButton("Deselect All")
        self.deselect_all_btn.setStyleSheet("""
            QPushButton {
                background: #6b7280;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover { background: #4b5563; }
        """)
        self.deselect_all_btn.clicked.connect(self._deselect_all)
        select_row.addWidget(self.deselect_all_btn)
        select_row.addStretch()
        panel2_layout.addLayout(select_row)

        # Track last clicked index for shift-select
        self._last_clicked_index = None

        # Results container
        self.awol_results_container = QVBoxLayout()
        self.awol_results_container.setSpacing(6)
        panel2_layout.addLayout(self.awol_results_container)

        scroll_layout.addWidget(panel2)
        scroll_layout.addStretch()
        scroll.setWidget(scroll_content)
        layout.addWidget(scroll, 1)

        add_lock_to_popup(self, show_button=False)

    def _on_yes(self):
        self.no_btn.setChecked(False)
        self.yes_btn.setChecked(True)
        self._send()

    def _on_no(self):
        self.yes_btn.setChecked(False)
        self.no_btn.setChecked(True)
        self._send()

    def set_notes(self, notes: list):
        """Search notes for AWOL episodes and display results."""
        import re
        from datetime import datetime

        # Clear previous results
        while self.awol_results_container.count():
            item = self.awol_results_container.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._awol_results = []

        if not notes:
            self.awol_count_label.setText("No notes available")
            return

        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        awol_lines = []
        for note in notes:
            content = note.get('content', '') or note.get('text', '') or ''
            content = content.replace('\r\n', '\n').replace('\r', '\n')

            date_val = note.get('date') or note.get('datetime')
            date_obj = parse_date(date_val)
            date_str = date_obj.strftime('%d/%m/%Y') if date_obj else 'Unknown'

            # Split into lines and sentences
            all_segments = []
            for line in content.split('\n'):
                sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', line)
                all_segments.extend(sentences)

            for line in all_segments:
                line_clean = line.strip()
                line_clean = ''.join(c for c in line_clean if ord(c) >= 32 or c == '\t')
                line_lower = line_clean.lower()

                if not line_clean or len(line_clean) < 10:
                    continue

                # Check for exclusion prefixes (Risk:, Risk to self:, etc.)
                if any(line_lower.startswith(prefix) for prefix in AWOL_EXCLUDE_PREFIXES):
                    continue

                # Check for exclusion terms
                if any(ex in line_lower for ex in AWOL_EXCLUDE_TERMS):
                    continue

                # Check for AWOL terms
                found_term = None
                term_position = -1
                for term in AWOL_SEARCH_TERMS:
                    pos = line_lower.find(term)
                    if pos >= 0:
                        found_term = term
                        term_position = pos
                        break

                if found_term and term_position >= 0:
                    # Check if negation word appears before the AWOL keyword
                    text_before_term = line_lower[:term_position]
                    negated = False
                    for neg_word in AWOL_NEGATION_WORDS:
                        # Check if negation word is within 30 chars before the term
                        # and is a whole word (not part of another word)
                        pattern = r'\b' + re.escape(neg_word) + r'\b'
                        if re.search(pattern, text_before_term[-30:] if len(text_before_term) > 30 else text_before_term):
                            negated = True
                            break

                    if not negated:
                        awol_lines.append({
                            "date": date_str,
                            "date_obj": date_obj,
                            "text": line_clean,
                            "term": found_term,
                        })

        # Sort by date (most recent first)
        awol_lines.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)

        # Remove duplicates (same text)
        seen_texts = set()
        unique_results = []
        for item in awol_lines:
            text_key = item["text"][:100].lower()
            if text_key not in seen_texts:
                seen_texts.add(text_key)
                unique_results.append(item)

        # Group by date - combine entries with same date
        date_groups = {}
        for item in unique_results:
            date_str = item["date"]
            if date_str not in date_groups:
                date_groups[date_str] = {
                    "date": date_str,
                    "date_obj": item["date_obj"],
                    "texts": [],
                    "terms": set(),
                }
            date_groups[date_str]["texts"].append(item["text"])
            date_groups[date_str]["terms"].add(item["term"])

        # Convert to list and combine texts
        grouped_results = []
        for date_str, group in date_groups.items():
            combined_text = " | ".join(group["texts"][:5])  # Limit to 5 entries per date
            if len(group["texts"]) > 5:
                combined_text += f" | ... (+{len(group['texts']) - 5} more)"
            grouped_results.append({
                "date": group["date"],
                "date_obj": group["date_obj"],
                "text": combined_text,
                "term": ", ".join(group["terms"]),
                "entry_count": len(group["texts"]),
            })

        # Sort grouped results by date (most recent first)
        grouped_results.sort(key=lambda x: x.get("date_obj") or datetime.min, reverse=True)

        self._awol_results = grouped_results

        if not grouped_results:
            self.awol_count_label.setText("No AWOL episodes found in notes")
        else:
            total_entries = sum(r["entry_count"] for r in grouped_results)
            self.awol_count_label.setText(f"Found {total_entries} AWOL mention(s) across {len(grouped_results)} date(s)")

            # Display results (max 20 dates)
            for i, result in enumerate(grouped_results[:20]):
                result_widget = self._create_result_widget(result, i)
                self.awol_results_container.addWidget(result_widget)

    def _create_result_widget(self, result, index):
        """Create a widget for a single AWOL search result."""
        from PySide6.QtWidgets import QCheckBox
        from functools import partial

        widget = QFrame()
        widget.setCursor(Qt.CursorShape.PointingHandCursor)
        widget.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
            }
        """)
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(8, 6, 8, 6)
        layout.setSpacing(8)

        # Checkbox to include in report
        cb = QCheckBox()
        cb.setChecked(False)
        cb.toggled.connect(lambda checked, idx=index: self._on_checkbox_toggled(idx, checked))
        layout.addWidget(cb)

        # Date
        date_label = QLabel(result["date"])
        date_label.setStyleSheet("font-weight: 600; color: #92400e; min-width: 70px;")
        layout.addWidget(date_label)

        # Text (truncated)
        text = result["text"]
        if len(text) > 150:
            text = text[:150] + "..."
        text_label = QLabel(text)
        text_label.setWordWrap(True)
        text_label.setStyleSheet("font-size: 16px; color: #374151;")
        layout.addWidget(text_label, 1)

        # Store references
        result["checkbox"] = cb
        result["widget"] = widget
        result["index"] = index

        # Make widget clickable
        widget.mousePressEvent = partial(self._on_widget_clicked, index)

        return widget

    def _on_widget_clicked(self, index, event):
        """Handle click on result widget - toggle checkbox with shift-select support."""
        from PySide6.QtCore import Qt
        from PySide6.QtWidgets import QApplication

        if index >= len(self._awol_results):
            return

        modifiers = QApplication.keyboardModifiers()
        result = self._awol_results[index]
        cb = result.get("checkbox")

        if modifiers & Qt.KeyboardModifier.ShiftModifier and self._last_clicked_index is not None:
            # Shift-click: select range
            start = min(self._last_clicked_index, index)
            end = max(self._last_clicked_index, index)
            for i in range(start, end + 1):
                if i < len(self._awol_results):
                    self._awol_results[i]["checkbox"].setChecked(True)
        else:
            # Normal click: toggle checkbox
            if cb:
                cb.setChecked(not cb.isChecked())

        self._last_clicked_index = index

    def _on_checkbox_toggled(self, index, checked):
        """Handle checkbox toggle - update styling and preview."""
        if index < len(self._awol_results):
            result = self._awol_results[index]
            widget = result.get("widget")
            if widget:
                if checked:
                    widget.setStyleSheet("""
                        QFrame {
                            background: #fef9c3;
                            border: 1px solid #facc15;
                            border-radius: 6px;
                        }
                    """)
                else:
                    widget.setStyleSheet("""
                        QFrame {
                            background: white;
                            border: 1px solid #d1d5db;
                            border-radius: 6px;
                        }
                    """)
            self._last_clicked_index = index

        # Auto-select "Yes" when any result is checked
        if checked:
            self.yes_btn.setChecked(True)
            self.no_btn.setChecked(False)
        self._send()

    def _select_all(self):
        """Select all AWOL checkboxes."""
        for result in self._awol_results:
            cb = result.get("checkbox")
            if cb:
                cb.setChecked(True)

    def _deselect_all(self):
        """Deselect all AWOL checkboxes."""
        for result in self._awol_results:
            cb = result.get("checkbox")
            if cb:
                cb.setChecked(False)

    def _get_selected_awol_results(self):
        """Get AWOL results that are checked."""
        selected = []
        for result in self._awol_results:
            cb = result.get("checkbox")
            if cb and cb.isChecked():
                selected.append(result)
        return selected

    def _formatted_text(self):
        parts = []

        if self.no_btn.isChecked():
            parts.append(f"There have been no occasions on which {self._pronoun.lower()} has been absent without leave or failed to return from granted leave.")
        elif self.yes_btn.isChecked():
            details = self.details_edit.toPlainText().strip()
            if details:
                parts.append(details)

            # Add selected AWOL results
            selected = self._get_selected_awol_results()
            if selected:
                awol_parts = []
                for result in selected:
                    awol_parts.append(f"[{result['date']}] {result['text']}")
                if awol_parts:
                    if parts:
                        parts.append("\n\nAWOL episodes from notes:")
                    else:
                        parts.append("AWOL episodes from notes:")
                    parts.extend(awol_parts)

            if not details and not selected:
                parts.append(f"{self._pronoun} has had AWOL concerns (details to be provided).")

        return "\n".join(parts)

    def _send(self):
        text = self._formatted_text()
        if text:
            self.sent.emit(text)

    def set_entries(self, entries: list, date_range_info: str = ""):
        """Set entries from data extractor - triggers AWOL search on entries."""
        if entries:
            self.set_notes(entries)
            print(f"[NURSING] AWOLPopup set_entries called with {len(entries)} entries")


class IncidentPopup(QWidget):
    """Popup for incident sections (12, 13, 14) with preview and checkboxes like AWOL."""

    sent = Signal(str)

    def __init__(self, title: str, section_type: str = "harm", parent=None, gender="Male"):
        super().__init__(parent)
        self._title = title
        self._section_type = section_type  # "harm", "property", or "seclusion"
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"
        self._incident_results = []
        self.notes = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # ========== SCROLLABLE CONTENT - INCIDENT RESULTS ==========
        results_frame = QFrame()
        results_frame.setStyleSheet("QFrame { background: #fef3c7; border: 1px solid #fcd34d; border-radius: 8px; }")
        results_layout = QVBoxLayout(results_frame)
        results_layout.setContentsMargins(12, 12, 12, 12)
        results_layout.setSpacing(8)

        header_text = {
            "harm": "Incidents of Harm to Self/Others",
            "property": "Incidents of Property Damage",
            "seclusion": "Seclusion/Restraint Episodes"
        }.get(section_type, "Incidents")

        results_header = QLabel(header_text)
        results_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #92400e;")
        results_layout.addWidget(results_header)

        self.count_label = QLabel("No data loaded yet")
        self.count_label.setStyleSheet("font-size: 17px; color: #92400e; font-weight: 600;")
        results_layout.addWidget(self.count_label)

        # Select All / Deselect All buttons
        select_row = QHBoxLayout()
        self.select_all_btn = QPushButton("Select All")
        self.select_all_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover { background: #d97706; }
        """)
        self.select_all_btn.clicked.connect(self._select_all)
        select_row.addWidget(self.select_all_btn)

        self.deselect_all_btn = QPushButton("Deselect All")
        self.deselect_all_btn.setStyleSheet("""
            QPushButton {
                background: #6b7280;
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 16px;
                font-weight: 600;
            }
            QPushButton:hover { background: #4b5563; }
        """)
        self.deselect_all_btn.clicked.connect(self._deselect_all)
        select_row.addWidget(self.deselect_all_btn)
        select_row.addStretch()
        results_layout.addLayout(select_row)

        # Track last clicked index for shift-select
        self._last_clicked_index = None

        # Scroll area for results
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.NoFrame)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        scroll_content = QWidget()
        self.results_container = QVBoxLayout(scroll_content)
        self.results_container.setSpacing(6)
        self.results_container.setContentsMargins(0, 0, 0, 0)
        scroll.setWidget(scroll_content)

        results_layout.addWidget(scroll, 1)
        layout.addWidget(results_frame, 1)

        add_lock_to_popup(self, show_button=False)

    def set_entries(self, entries: list, date_range_info: str = ""):
        """Set incident entries and display with checkboxes."""
        from datetime import datetime

        # Clear previous results
        while self.results_container.count():
            item = self.results_container.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self._incident_results = []

        if not entries:
            self.count_label.setText("No incidents found")
            return

        # Group entries by date
        date_groups = {}
        for entry in entries:
            date_str = entry.get('date', 'Unknown')
            content = entry.get('content', '') or entry.get('text', '')
            if not content.strip():
                continue

            if date_str not in date_groups:
                date_groups[date_str] = {
                    "date": date_str,
                    "texts": [],
                }
            # Avoid duplicates
            text_lower = content.strip().lower()[:100]
            if not any(t.lower()[:100] == text_lower for t in date_groups[date_str]["texts"]):
                date_groups[date_str]["texts"].append(content.strip())

        # Convert to list and combine texts
        grouped_results = []
        for date_str, group in date_groups.items():
            combined_text = " | ".join(group["texts"][:5])
            if len(group["texts"]) > 5:
                combined_text += f" | ... (+{len(group['texts']) - 5} more)"
            grouped_results.append({
                "date": group["date"],
                "text": combined_text,
                "entry_count": len(group["texts"]),
            })

        # Sort by date (try to parse dates for proper sorting)
        def parse_date_key(item):
            date_str = item.get("date", "")
            for fmt in ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(date_str[:10], fmt)
                except:
                    pass
            return datetime.min

        grouped_results.sort(key=parse_date_key, reverse=True)
        self._incident_results = grouped_results

        if not grouped_results:
            self.count_label.setText("No incidents found")
        else:
            total_entries = sum(r["entry_count"] for r in grouped_results)
            self.count_label.setText(f"Found {total_entries} incident(s) across {len(grouped_results)} date(s)")

            # Display all results with checkboxes (no limit)
            for i, result in enumerate(grouped_results):
                result_widget = self._create_result_widget(result, i)
                self.results_container.addWidget(result_widget)

        self.results_container.addStretch()

    def _create_result_widget(self, result, index):
        """Create a widget for a single incident result with checkbox."""
        from PySide6.QtWidgets import QCheckBox
        from functools import partial

        widget = QFrame()
        widget.setCursor(Qt.CursorShape.PointingHandCursor)
        widget.setStyleSheet("""
            QFrame {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
            }
        """)
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(8, 6, 8, 6)
        layout.setSpacing(8)

        # Checkbox to include in report
        cb = QCheckBox()
        cb.setChecked(False)
        cb.toggled.connect(lambda checked, idx=index: self._on_checkbox_toggled(idx, checked))
        layout.addWidget(cb)

        # Date
        date_label = QLabel(result["date"])
        date_label.setStyleSheet("font-weight: 600; color: #92400e; min-width: 70px;")
        layout.addWidget(date_label)

        # Text (truncated)
        text = result["text"]
        if len(text) > 200:
            text = text[:200] + "..."
        text_label = QLabel(text)
        text_label.setWordWrap(True)
        text_label.setStyleSheet("font-size: 16px; color: #374151;")
        layout.addWidget(text_label, 1)

        # Store references
        result["checkbox"] = cb
        result["widget"] = widget
        result["index"] = index

        # Make widget clickable
        widget.mousePressEvent = partial(self._on_widget_clicked, index)

        return widget

    def _on_widget_clicked(self, index, event):
        """Handle click on result widget - toggle checkbox with shift-select support."""
        from PySide6.QtCore import Qt
        from PySide6.QtWidgets import QApplication

        if index >= len(self._incident_results):
            return

        modifiers = QApplication.keyboardModifiers()
        result = self._incident_results[index]
        cb = result.get("checkbox")

        if modifiers & Qt.KeyboardModifier.ShiftModifier and self._last_clicked_index is not None:
            # Shift-click: select range
            start = min(self._last_clicked_index, index)
            end = max(self._last_clicked_index, index)
            for i in range(start, end + 1):
                if i < len(self._incident_results):
                    self._incident_results[i]["checkbox"].setChecked(True)
        else:
            # Normal click: toggle checkbox
            if cb:
                cb.setChecked(not cb.isChecked())

        self._last_clicked_index = index

    def _on_checkbox_toggled(self, index, checked):
        """Handle checkbox toggle - update styling and send."""
        if index < len(self._incident_results):
            result = self._incident_results[index]
            widget = result.get("widget")
            if widget:
                if checked:
                    widget.setStyleSheet("""
                        QFrame {
                            background: #fef9c3;
                            border: 1px solid #facc15;
                            border-radius: 6px;
                        }
                    """)
                else:
                    widget.setStyleSheet("""
                        QFrame {
                            background: white;
                            border: 1px solid #d1d5db;
                            border-radius: 6px;
                        }
                    """)
            self._last_clicked_index = index
        self._send()

    def _select_all(self):
        """Select all incident checkboxes."""
        for result in self._incident_results:
            cb = result.get("checkbox")
            if cb:
                cb.setChecked(True)

    def _deselect_all(self):
        """Deselect all incident checkboxes."""
        for result in self._incident_results:
            cb = result.get("checkbox")
            if cb:
                cb.setChecked(False)

    def _get_selected_results(self):
        """Get results that are checked."""
        selected = []
        for result in self._incident_results:
            cb = result.get("checkbox")
            if cb and cb.isChecked():
                selected.append(result)
        return selected

    def _formatted_text(self):
        """Generate formatted text for preview/send."""
        selected = self._get_selected_results()

        if not selected:
            # Default text based on section type
            if self._section_type == "seclusion":
                return f"There have been no occasions on which {self._pronoun.lower()} has been secluded or restrained during this admission."
            elif self._section_type == "property":
                return "There have been no concerns regarding property damage."
            else:
                return "There have been no concerns regarding harm to self or others."

        # Build incident list
        parts = []
        for result in selected:
            parts.append(f"[{result['date']}] {result['text']}")

        return "\n\n".join(parts)

    def _send(self):
        """Send the formatted text to the report."""
        text = self._formatted_text()
        if text:
            self.sent.emit(text)


class SeclusionRestraintPopup(QWidget):
    """Popup for seclusion/restraint section with Yes/No radio and conditional explanation box."""

    sent = Signal(str)

    def __init__(self, parent=None, gender="Male"):
        super().__init__(parent)
        self._gender = gender
        self._pronoun = "He" if gender == "Male" else "She"

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Have there been any occasions on which the patient has been secluded or restrained?")
        label.setWordWrap(True)
        label.setStyleSheet("font-weight: 600; color: #374151;")
        layout.addWidget(label)

        # Yes/No radio buttons
        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes")
        self.no_btn = QPushButton("No")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6;
                    border: 1px solid #d1d5db;
                    padding: 8px 24px;
                    border-radius: 6px;
                }
                QPushButton:checked {
                    background: #14b8a6;
                    color: white;
                    border-color: #14b8a6;
                }
            """)
        self.yes_btn.clicked.connect(self._on_yes_clicked)
        self.no_btn.clicked.connect(self._on_no_clicked)
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        # Reason container (hidden by default, shown when Yes clicked)
        self.reason_container = QWidget()
        reason_layout = QVBoxLayout(self.reason_container)
        reason_layout.setContentsMargins(0, 8, 0, 0)
        reason_layout.setSpacing(8)

        reason_label = QLabel("Please explain why seclusion or restraint was necessary:")
        reason_label.setStyleSheet("font-weight: 600; color: #374151;")
        reason_layout.addWidget(reason_label)

        self.reason_edit = QTextEdit()
        self.reason_edit.setPlaceholderText("Explain why necessary...")
        self.reason_edit.setMinimumHeight(120)
        reason_layout.addWidget(self.reason_edit)

        self.reason_container.setVisible(False)  # Hidden by default
        layout.addWidget(self.reason_container)

        layout.addStretch()

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #14b8a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #0d9488;
            }
        """)
        send_btn.clicked.connect(self._send)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)

    def _on_yes_clicked(self):
        self.no_btn.setChecked(False)
        self.reason_container.setVisible(True)

    def _on_no_clicked(self):
        self.yes_btn.setChecked(False)
        self.reason_container.setVisible(False)

    def _send(self):
        if self.no_btn.isChecked():
            # Output "No" prefix so export knows to tick No checkbox
            self.sent.emit("No - There have been no occasions of seclusion or restraint during this admission.")
        elif self.yes_btn.isChecked():
            reason = self.reason_edit.toPlainText().strip()
            if reason:
                # Output "Yes" prefix so export knows to tick Yes checkbox
                self.sent.emit(f"Yes - {reason}")
            else:
                self.sent.emit("Yes - Seclusion or restraint was necessary (details to be provided).")


class OtherInfoPopup(QWidget):
    """Popup for other relevant information section."""

    sent = Signal(str)

    def __init__(self, parent=None):
        super().__init__(parent)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        label = QLabel("Is there any other relevant information that the tribunal should know?")
        label.setWordWrap(True)
        label.setStyleSheet("font-size: 22px; font-weight: 600; color: #374151;")
        layout.addWidget(label)

        # Yes/No
        btn_row = QHBoxLayout()
        self.yes_btn = QPushButton("Yes")
        self.no_btn = QPushButton("No")
        for btn in [self.yes_btn, self.no_btn]:
            btn.setCheckable(True)
            btn.setStyleSheet("""
                QPushButton {
                    background: #f3f4f6;
                    border: 1px solid #d1d5db;
                    padding: 8px 24px;
                    border-radius: 6px;
                    font-size: 22px;
                }
                QPushButton:checked {
                    background: #14b8a6;
                    color: white;
                    border-color: #14b8a6;
                }
            """)
        self.yes_btn.clicked.connect(lambda: self.no_btn.setChecked(False))
        self.no_btn.clicked.connect(lambda: self.yes_btn.setChecked(False))
        btn_row.addWidget(self.yes_btn)
        btn_row.addWidget(self.no_btn)
        btn_row.addStretch()
        layout.addLayout(btn_row)

        # Details
        self.details_edit = QTextEdit()
        self.details_edit.setStyleSheet("font-size: 22px;")
        self.details_edit.setPlaceholderText("Provide other relevant information...")
        self.details_edit.setMinimumHeight(120)
        layout.addWidget(self.details_edit)

        layout.addStretch()

        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #14b8a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #0d9488;
            }
        """)
        send_btn.clicked.connect(self._send)
        layout.addWidget(send_btn)

        add_lock_to_popup(self, show_button=False)

    def _send(self):
        if self.no_btn.isChecked():
            self.sent.emit("There is no other relevant information for the tribunal.")
        elif self.yes_btn.isChecked():
            details = self.details_edit.toPlainText().strip()
            if details:
                self.sent.emit(details)
