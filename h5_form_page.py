# ================================================================
#  H5 FORM PAGE — Renewal of Authority for Detention
#  Mental Health Act 1983 - Form H5 Regulation 13(1), (2) and (3)
#  Section 20 — Renewal of authority for detention
#  CARD/POPUP LAYOUT with ResizableSection
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QSpinBox, QRadioButton, QButtonGroup, QComboBox, QCompleter,
    QStyleFactory, QSlider, QStackedWidget, QSplitter
)
from utils.resource_path import resource_path


# ================================================================
# RESIZABLE SECTION WITH DRAG BAR
# ================================================================
class ResizableSection(QFrame):
    """Section with a draggable bottom edge to resize height."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._min_height = 60
        self._max_height = 400
        self._content = None
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0

        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(0)

        # Content container
        self._content_container = QWidget()
        self._content_layout = QVBoxLayout(self._content_container)
        self._content_layout.setContentsMargins(12, 8, 12, 8)
        self._layout.addWidget(self._content_container, 1)

        # Drag handle at bottom
        self._drag_handle = QFrame()
        self._drag_handle.setFixedHeight(8)
        self._drag_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        self._drag_handle.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
                border-radius: 2px;
            }
            QFrame:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.3 #dc2626, stop:0.7 #dc2626, stop:1 transparent);
            }
        """)
        self._drag_handle.mousePressEvent = self._handle_press
        self._drag_handle.mouseMoveEvent = self._handle_move
        self._drag_handle.mouseReleaseEvent = self._handle_release
        self._layout.addWidget(self._drag_handle)

    def set_content(self, widget):
        self._content = widget
        self._content_layout.addWidget(widget)

    def set_content_height(self, h):
        h = max(self._min_height, min(self._max_height, h))
        self._content_container.setFixedHeight(h)

    def _handle_press(self, event):
        self._dragging = True
        self._drag_start_y = event.globalPosition().y()
        self._drag_start_height = self._content_container.height()

    def _handle_move(self, event):
        if self._dragging:
            delta = event.globalPosition().y() - self._drag_start_y
            new_height = self._drag_start_height + delta
            new_height = max(self._min_height, min(self._max_height, new_height))
            self._content_container.setFixedHeight(int(new_height))

    def _handle_release(self, event):
        self._dragging = False

# ICD-10 data - curated list matching iOS app
try:
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []

from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# NO-WHEEL COMBOBOX (prevents scroll from changing value)
# ================================================================
class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelSpinBox(QSpinBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# H5 CARD WIDGET - Fixed header, scrollable content
# ================================================================
class H5CardWidget(QFrame):
    """Card with fixed header and scrollable content area."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("h5Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#h5Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#h5Card:hover {
                border-color: #dc2626;
                background: #fef2f2;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QLabel#cardTitle {
                font-size: 20px;
                font-weight: 600;
                color: #dc2626;
                padding-bottom: 4px;
            }
            QFrame#divider {
                background: #e5e7eb;
                height: 1px;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 12)
        layout.setSpacing(0)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setObjectName("cardTitle")
        self.title_label.setFixedHeight(24)
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        # Divider
        divider = QFrame()
        divider.setObjectName("divider")
        divider.setFixedHeight(1)
        layout.addWidget(divider)

        # Scrollable content
        self.content_scroll = QScrollArea()
        self.content_scroll.setWidgetResizable(True)
        self.content_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        self.content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.content_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        content_widget = QWidget()
        content_widget.setStyleSheet("background: transparent;")
        self.content_layout = QVBoxLayout(content_widget)
        self.content_layout.setContentsMargins(0, 8, 0, 4)
        self.content_layout.setSpacing(0)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 20px;
                color: #374151;
                padding: 4px;
                background: transparent;
                border: none;
            }
        """)
        self.content.setFrameShape(QFrame.Shape.NoFrame)
        self.content_layout.addWidget(self.content)
        self.content_layout.addStretch()

        self.content_scroll.setWidget(content_widget)
        layout.addWidget(self.content_scroll, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_active(self, active: bool):
        """Set active state - shows persistent highlight."""
        if active:
            self.setStyleSheet("""
                QFrame#h5Card {
                    background: #fef2f2;
                    border: 2px solid #dc2626;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
                QLabel#cardTitle {
                    font-size: 20px;
                    font-weight: 600;
                    color: #dc2626;
                }
                QFrame#divider {
                    background: #e5e7eb;
                    height: 1px;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#h5Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#h5Card:hover {
                    border-color: #dc2626;
                    background: #fef2f2;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
                QLabel#cardTitle {
                    font-size: 20px;
                    font-weight: 600;
                    color: #dc2626;
                }
                QFrame#divider {
                    background: #e5e7eb;
                    height: 1px;
                }
            """)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# H5 TOOLBAR
# ================================================================
class H5Toolbar(QWidget):
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            H5Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN H5 FORM PAGE - Card/Popup Layout
# ================================================================
class H5FormPage(QWidget):
    """Page for completing MHA Form H5 - Renewal of Authority for Detention."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race", "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}

        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {"full_name": details[1] or "", "email": details[7] or ""}

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #dc2626; border-bottom: 1px solid #b91c1c;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 21px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form H5 — Renewal of Authority for Detention")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - cards on left, popup on right with splitter
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(16, 16, 16, 16)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #dc2626, stop:0.7 #dc2626, stop:1 transparent);
            }
        """)

        # Left: Cards column (scrollable)
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(0, 0, 8, 0)
        self.cards_layout.setSpacing(8)

        # Create cards with ResizableSection
        self._create_details_card()  # Combined: Patient, Hospital, Clinician
        self._create_clinical_card()
        self._create_informal_card()
        self._create_signature_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        self.main_splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)

        # Create popup panels
        self._create_details_popup()  # Combined: Patient, Hospital, Clinician
        self._create_clinical_popup()
        self._create_informal_popup()
        self._create_signature_popup()

        # Initialize cards with default date values
        self._update_details_card()
        self._update_signature_card()

        self.main_splitter.addWidget(self.popup_stack)
        self.main_splitter.setSizes([280, 600])  # Initial sizes
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        # Connect live preview updates
        self._connect_clinical_live_preview()
        self._connect_informal_live_preview()

        # Show first popup by default
        self._on_card_clicked("details")

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"details": 0, "clinical": 1, "informal": 2, "signatures": 3}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            # Highlight active card
            for k, card in self.cards.items():
                card.set_active(k == key)
            # Sync signature displays when opening signature popup
            if key == "signatures":
                self._sync_signature_displays()

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_details_card(self):
        """Combined card for Patient, Hospital & Dates, Clinician."""
        section = ResizableSection()
        section.set_content_height(220)
        section._min_height = 150
        section._max_height = 400

        card = H5CardWidget("Patient / Hospital / Clinician", "details")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["details"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_clinical_card(self):
        section = ResizableSection()
        section.set_content_height(250)
        section._min_height = 150
        section._max_height = 450

        card = H5CardWidget("Reasons for Renewal", "clinical")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["clinical"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_informal_card(self):
        section = ResizableSection()
        section.set_content_height(200)
        section._min_height = 120
        section._max_height = 350

        card = H5CardWidget("Why Detention Required", "informal")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["informal"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signature_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 300

        card = H5CardWidget("Signatures", "signatures")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signatures"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # HELPER METHODS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 21px;
            }
            QLineEdit:focus { border-color: #dc2626; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 21px;
            }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        return date_edit

    def _create_styled_frame(self, color: str) -> QFrame:
        colors = {
            "red": ("#fef2f2", "#fecaca"),
            "green": ("#f0fdf4", "#bbf7d0"),
            "blue": ("#eff6ff", "#bfdbfe"),
            "yellow": ("#fefce8", "#fef08a"),
            "purple": ("#faf5ff", "#e9d5ff"),
        }
        bg, border = colors.get(color, ("#f9fafb", "#e5e7eb"))
        frame = QFrame()
        frame.setStyleSheet(f"QFrame {{ background: {bg}; border: none; border-radius: 8px; }}")
        return frame

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_details_popup(self):
        """Combined popup for Patient, Hospital & Dates, Clinician."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        header = QLabel("Patient / Hospital / Clinician")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # --- Patient Section ---
        patient_frame = self._create_styled_frame("red")
        pf_layout = QVBoxLayout(patient_frame)
        pf_layout.setContentsMargins(12, 10, 12, 10)
        pf_layout.setSpacing(8)

        name_lbl = QLabel("Patient Name")
        name_lbl.setStyleSheet("font-size: 19px; font-weight: 600; color: #991b1b;")
        pf_layout.addWidget(name_lbl)
        self.patient_name = self._create_line_edit("Full name")
        pf_layout.addWidget(self.patient_name)

        # Demographics row
        demo_row = QHBoxLayout()
        demo_row.setSpacing(8)

        # Age - visible for clinical reasons text generation
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 19px; color: #374151;")
        demo_row.addWidget(age_lbl)
        self.age_spin = NoWheelSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setFixedWidth(55)
        self.age_spin.setStyleSheet("font-size: 19px; padding: 4px;")
        demo_row.addWidget(self.age_spin)

        # Gender
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("O")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 19px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 18px;
                    height: 18px;
                }
            """)
            self.gender_group.addButton(rb)
            demo_row.addWidget(rb)

        # Ethnicity - visible for clinical reasons text generation
        eth_lbl = QLabel("Ethnicity:")
        eth_lbl.setStyleSheet("font-size: 19px; color: #374151;")
        demo_row.addWidget(eth_lbl)
        self.ethnicity_combo = NoWheelComboBox()
        self.ethnicity_combo.addItem("Not specified")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(140)
        self.ethnicity_combo.setStyleSheet("font-size: 17px; padding: 4px;")
        demo_row.addWidget(self.ethnicity_combo)
        demo_row.addStretch()

        pf_layout.addLayout(demo_row)
        layout.addWidget(patient_frame)

        # --- Hospital Section ---
        hosp_frame = self._create_styled_frame("blue")
        hf_layout = QVBoxLayout(hosp_frame)
        hf_layout.setContentsMargins(12, 10, 12, 10)
        hf_layout.setSpacing(8)

        hosp_lbl = QLabel("To managers of:")
        hosp_lbl.setStyleSheet("font-size: 19px; font-weight: 600; color: #1e40af;")
        hf_layout.addWidget(hosp_lbl)
        self.hospital = self._create_line_edit("Hospital name and address")
        hf_layout.addWidget(self.hospital)

        dates_row = QHBoxLayout()
        dates_row.setSpacing(12)

        exam_lbl = QLabel("Examined:")
        exam_lbl.setStyleSheet("font-size: 19px; font-weight: 500; color: #374151;")
        dates_row.addWidget(exam_lbl)
        self.exam_date = self._create_date_edit()
        self.exam_date.setFixedWidth(160)
        dates_row.addWidget(self.exam_date)

        exp_lbl = QLabel("Expires:")
        exp_lbl.setStyleSheet("font-size: 19px; font-weight: 500; color: #374151;")
        dates_row.addWidget(exp_lbl)
        self.expiry_date = self._create_date_edit()
        self.expiry_date.setFixedWidth(160)
        dates_row.addWidget(self.expiry_date)
        dates_row.addStretch()

        hf_layout.addLayout(dates_row)
        layout.addWidget(hosp_frame)

        # --- Clinician Section ---
        rc_frame = self._create_styled_frame("green")
        rc_layout = QVBoxLayout(rc_frame)
        rc_layout.setContentsMargins(12, 10, 12, 10)
        rc_layout.setSpacing(6)

        rc_header = QLabel("Responsible Clinician")
        rc_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #166534;")
        rc_layout.addWidget(rc_header)

        rc_row = QHBoxLayout()
        rc_row.setSpacing(8)
        self.rc_name = self._create_line_edit("RC full name")
        rc_row.addWidget(self.rc_name, 1)
        self.rc_profession = self._create_line_edit("Profession")
        self.rc_profession.setText("Consultant Psychiatrist")
        rc_row.addWidget(self.rc_profession, 1)
        rc_layout.addLayout(rc_row)

        rc_row2 = QHBoxLayout()
        rc_row2.setSpacing(8)
        self.rc_address = self._create_line_edit("Address")
        rc_row2.addWidget(self.rc_address, 2)
        self.rc_email = self._create_line_edit("Email")
        rc_row2.addWidget(self.rc_email, 1)
        rc_layout.addLayout(rc_row2)

        layout.addWidget(rc_frame)

        # --- Consulted Section ---
        consult_frame = self._create_styled_frame("purple")
        cf_layout = QVBoxLayout(consult_frame)
        cf_layout.setContentsMargins(12, 10, 12, 10)
        cf_layout.setSpacing(6)

        consult_header = QLabel("Professional Consulted")
        consult_header.setStyleSheet("font-size: 19px; font-weight: 700; color: #7c3aed;")
        cf_layout.addWidget(consult_header)

        self.consulted_name = self._create_line_edit("Full name")
        cf_layout.addWidget(self.consulted_name)

        self.consulted_profession = NoWheelComboBox()
        self.consulted_profession.addItem("Select profession...")
        self.consulted_profession.addItems([
            "Registered Mental Health Nurse",
            "Registered Learning Disabilities Nurse",
            "Occupational Therapist",
            "Social Worker",
            "Psychologist"
        ])
        self.consulted_profession.setStyleSheet("QComboBox { padding: 8px; border: 1px solid #d1d5db; border-radius: 6px; font-size: 20px; }")
        cf_layout.addWidget(self.consulted_profession)

        layout.addWidget(consult_frame)

        # Auto-sync to card
        self.patient_name.textChanged.connect(self._update_details_card)
        self.hospital.textChanged.connect(self._update_details_card)
        self.exam_date.dateChanged.connect(self._update_details_card)
        self.expiry_date.dateChanged.connect(self._update_details_card)
        self.rc_name.textChanged.connect(self._update_details_card)
        self.rc_profession.textChanged.connect(self._update_details_card)
        self.rc_address.textChanged.connect(self._update_details_card)
        self.rc_email.textChanged.connect(self._update_details_card)
        self.consulted_name.textChanged.connect(self._update_details_card)
        self.consulted_profession.currentIndexChanged.connect(self._update_details_card)

        layout.addStretch()

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_clinical_popup(self):
        # Main container - NOT a scroll area, so preview stays fixed
        popup = QWidget()
        popup.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(popup)
        main_layout.setContentsMargins(20, 16, 20, 16)
        main_layout.setSpacing(12)

        header = QLabel("Reasons for Renewal")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        main_layout.addWidget(header)

        # Hidden label for state storage (used in export and state management)
        self.clinical_preview_label = QLabel("")
        self.clinical_preview_label.hide()

        # Scrollable form content
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        form_container = QWidget()
        form_container.setStyleSheet("background: transparent;")
        layout = QVBoxLayout(form_container)
        layout.setContentsMargins(0, 0, 8, 0)
        layout.setSpacing(12)

        # --- Mental Disorder (ICD-10) ---
        md_frame = self._create_styled_frame("green")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(6)

        md_header = QLabel("Mental Disorder (ICD-10)")
        md_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        # Primary diagnosis dropdown with grouped items
        self.dx_primary = NoWheelComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select primary diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.setMaximumWidth(400)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        md_layout.addWidget(self.dx_primary)

        # Secondary diagnosis dropdown
        self.dx_secondary = NoWheelComboBox()
        self.dx_secondary.setEditable(True)
        self.dx_secondary.lineEdit().setPlaceholderText("Secondary diagnosis (optional)...")
        self.dx_secondary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_secondary.setMaximumWidth(400)
        self.dx_secondary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_secondary.addItem(f"── {group_name} ──", None)
            idx = self.dx_secondary.count() - 1
            self.dx_secondary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_secondary.addItem(dx, dx)
        completer2 = QCompleter(ICD10_FLAT, self.dx_secondary)
        completer2.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer2.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_secondary.setCompleter(completer2)
        self.dx_secondary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        md_layout.addWidget(self.dx_secondary)

        layout.addWidget(md_frame)

        # --- Legal Criteria ---
        lc_frame = self._create_styled_frame("blue")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(4)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 0, 0, 0)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.chronic_cb = QCheckBox("Chronic and enduring")
        for cb in [self.relapsing_cb, self.treatment_resistant_cb, self.chronic_cb]:
            cb.setStyleSheet("font-size: 19px; color: #6b7280;")
            nature_opt_layout.addWidget(cb)
        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 0, 0, 0)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setFixedWidth(100)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)
        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 19px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 19px; padding: 6px; border: 1px solid #d1d5db; border-radius: 4px;")
        degree_opt_layout.addWidget(self.degree_details)
        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health
        self.nec_health = QCheckBox("Health")
        self.nec_health.setStyleSheet("font-size: 19px; color: #374151;")
        self.nec_health.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.nec_health)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 19px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)
        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 18px; color: #9ca3af;")
        mh_opt_layout.addWidget(self.poor_compliance_cb)
        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 18px; color: #9ca3af;")
        mh_opt_layout.addWidget(self.limited_insight_cb)
        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 19px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 18px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety
        self.nec_safety = QCheckBox("Safety")
        self.nec_safety.setStyleSheet("font-size: 19px; color: #374151;")
        self.nec_safety.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.nec_safety)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(2)

        # Self
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 19px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 18px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)
        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_risky = QCheckBox("Risky situations")
        self.self_hist_harm = QCheckBox("Self harm")
        for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm]:
            cb.setStyleSheet("font-size: 18px; color: #9ca3af;")
            self_opt_layout.addWidget(cb)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 18px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        self_opt_layout.addWidget(self_curr_lbl)
        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_risky = QCheckBox("Risky situations")
        self.self_curr_harm = QCheckBox("Self harm")
        for cb in [self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
            cb.setStyleSheet("font-size: 18px; color: #9ca3af;")
            self_opt_layout.addWidget(cb)
        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others
        self.nec_others = QCheckBox("Others")
        self.nec_others.setStyleSheet("font-size: 19px; font-weight: 600; color: #6b7280;")
        self.nec_others.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.nec_others)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 18px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)
        self.others_hist_violence = QCheckBox("Violence")
        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_arson = QCheckBox("Arson")
        for cb in [self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual, self.others_hist_stalking, self.others_hist_arson]:
            cb.setStyleSheet("font-size: 18px; color: #9ca3af;")
            others_opt_layout.addWidget(cb)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 18px; color: #9ca3af; font-weight: 600; margin-top: 2px;")
        others_opt_layout.addWidget(others_curr_lbl)
        self.others_curr_violence = QCheckBox("Violence")
        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_arson = QCheckBox("Arson")
        for cb in [self.others_curr_violence, self.others_curr_verbal, self.others_curr_sexual, self.others_curr_stalking, self.others_curr_arson]:
            cb.setStyleSheet("font-size: 18px; color: #9ca3af;")
            others_opt_layout.addWidget(cb)
        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        layout.addWidget(lc_frame)
        layout.addStretch()

        # Add form container to scroll area, then scroll area to main layout
        form_scroll.setWidget(form_container)
        main_layout.addWidget(form_scroll, 1)  # stretch factor 1 so it takes remaining space

        self.popup_stack.addWidget(popup)

    def _create_informal_popup(self):
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        header = QLabel("Why Detention Required")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Hidden label for state storage (preview removed - card auto-syncs)
        self.informal_preview_label = QLabel("")
        self.informal_preview_label.hide()

        # Options frame - all in one consistent background
        inf_frame = self._create_styled_frame("red")
        inf_layout = QVBoxLayout(inf_frame)
        inf_layout.setContentsMargins(16, 12, 16, 12)
        inf_layout.setSpacing(8)

        info = QLabel("Treatment cannot be provided unless the patient continues to be detained.")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 19px; color: #6b7280;")
        inf_layout.addWidget(info)

        inf_header = QLabel("Why Informal Not Appropriate")
        inf_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #991b1b; margin-top: 8px;")
        inf_layout.addWidget(inf_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed informal")
        self.insight_cb = QCheckBox("Lack of Insight")
        self.compliance_cb = QCheckBox("Compliance Issues")
        self.supervision_cb = QCheckBox("Needs MHA Supervision")
        for cb in [self.tried_failed_cb, self.insight_cb, self.compliance_cb, self.supervision_cb]:
            cb.setStyleSheet("font-size: 20px; color: #374151;")
            inf_layout.addWidget(cb)

        layout.addWidget(inf_frame)
        layout.addStretch()

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_signature_popup(self):
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        header = QLabel("Signatures")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Part 2: Professional Agreement
        prof_header = QLabel("Part 2: Professional Agreement")
        prof_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #7c3aed;")
        layout.addWidget(prof_header)

        prof_info = QLabel("I agree with the responsible clinician that this patient meets the criteria for renewal.")
        prof_info.setWordWrap(True)
        prof_info.setStyleSheet("font-size: 19px; color: #6b7280;")
        layout.addWidget(prof_info)

        self.sig_consulted_display = QLabel("Consulted professional: (set in Details panel)")
        self.sig_consulted_display.setStyleSheet("font-size: 20px; color: #374151; font-style: italic;")
        layout.addWidget(self.sig_consulted_display)

        row2 = QHBoxLayout()
        row2.setSpacing(12)
        sig_lbl = QLabel("Date:")
        sig_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        row2.addWidget(sig_lbl)
        self.prof_sig_date = self._create_date_edit()
        self.prof_sig_date.setFixedWidth(160)
        row2.addWidget(self.prof_sig_date)
        row2.addStretch()
        layout.addLayout(row2)

        # Divider
        divider = QFrame()
        divider.setFixedHeight(1)
        divider.setStyleSheet("background: #e5e7eb; margin: 12px 0;")
        layout.addWidget(divider)

        # Part 3: RC Signature
        rc_header = QLabel("Part 3: RC Signature")
        rc_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #166534;")
        layout.addWidget(rc_header)

        rc_name = self._my_details.get("full_name", "")
        self.sig_rc_display = QLabel(f"RC: {rc_name}" if rc_name else "RC: (set in Details panel)")
        self.sig_rc_display.setStyleSheet("font-size: 20px; color: #374151; font-weight: 600;")
        layout.addWidget(self.sig_rc_display)

        furnish_lbl = QLabel("I am furnishing this report by:")
        furnish_lbl.setStyleSheet("font-size: 19px; font-weight: 600; color: #374151;")
        layout.addWidget(furnish_lbl)

        self.furnish_group = QButtonGroup(self)
        self.furnish_internal = QRadioButton("Internal mail system")
        self.furnish_electronic = QRadioButton("Electronic communication")
        self.furnish_other = QRadioButton("Other delivery method")
        for rb in [self.furnish_internal, self.furnish_electronic, self.furnish_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 19px;
                    color: #374151;
                    margin-left: 8px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 18px;
                    height: 18px;
                }
            """)
            self.furnish_group.addButton(rb)
            layout.addWidget(rb)
        self.furnish_internal.setChecked(True)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(12)
        sig_lbl2 = QLabel("Signature Date:")
        sig_lbl2.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl2)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(160)
        sig_row.addWidget(self.rc_sig_date)
        sig_row.addStretch()
        layout.addLayout(sig_row)

        # Auto-sync to card
        self.furnish_internal.toggled.connect(self._update_signature_card)
        self.furnish_electronic.toggled.connect(self._update_signature_card)
        self.furnish_other.toggled.connect(self._update_signature_card)
        self.rc_sig_date.dateChanged.connect(self._update_signature_card)

        layout.addStretch()

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # TOGGLE HANDLERS
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.nec_others.setChecked(False)

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm,
                       self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
                cb.setChecked(False)

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            for cb in [self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual,
                       self.others_hist_stalking, self.others_hist_arson, self.others_curr_violence,
                       self.others_curr_verbal, self.others_curr_sexual, self.others_curr_stalking,
                       self.others_curr_arson]:
                cb.setChecked(False)

    # ----------------------------------------------------------------
    # LIVE PREVIEW METHODS
    # ----------------------------------------------------------------
    def _connect_clinical_live_preview(self):
        """Connect all clinical popup controls to live preview updates."""
        # Diagnosis text fields
        self.dx_primary.currentTextChanged.connect(self._update_clinical_preview)
        self.dx_secondary.currentTextChanged.connect(self._update_clinical_preview)

        # Nature checkboxes
        self.nature_cb.toggled.connect(self._update_clinical_preview)
        self.relapsing_cb.toggled.connect(self._update_clinical_preview)
        self.treatment_resistant_cb.toggled.connect(self._update_clinical_preview)
        self.chronic_cb.toggled.connect(self._update_clinical_preview)

        # Degree controls
        self.degree_cb.toggled.connect(self._update_clinical_preview)
        self.degree_slider.valueChanged.connect(self._update_clinical_preview)
        self.degree_details.textChanged.connect(self._update_clinical_preview)

        # Health checkboxes
        self.nec_health.toggled.connect(self._update_clinical_preview)
        self.mental_health_cb.toggled.connect(self._update_clinical_preview)
        self.poor_compliance_cb.toggled.connect(self._update_clinical_preview)
        self.limited_insight_cb.toggled.connect(self._update_clinical_preview)
        self.physical_health_cb.toggled.connect(self._update_clinical_preview)
        self.physical_health_details.textChanged.connect(self._update_clinical_preview)

        # Safety checkboxes
        self.nec_safety.toggled.connect(self._update_clinical_preview)
        self.self_harm_cb.toggled.connect(self._update_clinical_preview)
        self.self_hist_neglect.toggled.connect(self._update_clinical_preview)
        self.self_hist_risky.toggled.connect(self._update_clinical_preview)
        self.self_hist_harm.toggled.connect(self._update_clinical_preview)
        self.self_curr_neglect.toggled.connect(self._update_clinical_preview)
        self.self_curr_risky.toggled.connect(self._update_clinical_preview)
        self.self_curr_harm.toggled.connect(self._update_clinical_preview)

        # Others checkboxes
        self.nec_others.toggled.connect(self._update_clinical_preview)
        self.others_hist_violence.toggled.connect(self._update_clinical_preview)
        self.others_hist_verbal.toggled.connect(self._update_clinical_preview)
        self.others_hist_sexual.toggled.connect(self._update_clinical_preview)
        self.others_hist_stalking.toggled.connect(self._update_clinical_preview)
        self.others_hist_arson.toggled.connect(self._update_clinical_preview)
        self.others_curr_violence.toggled.connect(self._update_clinical_preview)
        self.others_curr_verbal.toggled.connect(self._update_clinical_preview)
        self.others_curr_sexual.toggled.connect(self._update_clinical_preview)
        self.others_curr_stalking.toggled.connect(self._update_clinical_preview)
        self.others_curr_arson.toggled.connect(self._update_clinical_preview)

        # Patient info that affects preview text
        self.patient_name.textChanged.connect(self._update_clinical_preview)
        self.age_spin.valueChanged.connect(self._update_clinical_preview)
        self.gender_male.toggled.connect(self._update_clinical_preview)
        self.gender_female.toggled.connect(self._update_clinical_preview)
        self.gender_other.toggled.connect(self._update_clinical_preview)
        self.ethnicity_combo.currentIndexChanged.connect(self._update_clinical_preview)

    def _connect_informal_live_preview(self):
        """Connect all informal popup controls to live preview updates."""
        self.tried_failed_cb.toggled.connect(self._update_informal_preview)
        self.insight_cb.toggled.connect(self._update_informal_preview)
        self.compliance_cb.toggled.connect(self._update_informal_preview)
        self.supervision_cb.toggled.connect(self._update_informal_preview)

        # Patient info that affects preview text
        self.patient_name.textChanged.connect(self._update_informal_preview)
        self.gender_male.toggled.connect(self._update_informal_preview)
        self.gender_female.toggled.connect(self._update_informal_preview)
        self.gender_other.toggled.connect(self._update_informal_preview)

    def _update_clinical_preview(self):
        """Update the clinical preview label and auto-sync to card."""
        text = self._generate_reasons_text()
        self.clinical_preview_label.setText(text if text else "Select options above to generate preview...")
        # Auto-sync to card
        self.cards["clinical"].set_content_text(text if text else "")

    def _update_informal_preview(self):
        """Update the informal preview label and auto-sync to card."""
        text = self._generate_informal_text()
        self.informal_preview_label.setText(text if text else "Select options above to generate preview...")
        # Auto-sync to card
        self.cards["informal"].set_content_text(text if text else "")

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_details_card(self):
        """Update combined details card with patient name, hospital, clinician info."""
        parts = []
        # Patient name only
        if self.patient_name.text():
            parts.append(self.patient_name.text())
        # Hospital name only
        if self.hospital.text():
            parts.append(self.hospital.text())
        # RC info
        if self.rc_name.text():
            parts.append(self.rc_name.text())
        # Consulted info
        if self.consulted_name.text():
            parts.append(self.consulted_name.text())
        self.cards["details"].set_content_text("\n".join(parts) if parts else "Click to enter details")
        # Update signature popup displays
        self._sync_signature_displays()

    def _sync_signature_displays(self):
        """Sync signature popup displays with clinician/consulted info."""
        # Update consulted display
        consulted = self.consulted_name.text()
        prof = self.consulted_profession.currentText() if self.consulted_profession.currentIndex() > 0 else ""
        if consulted:
            display_text = consulted
            if prof:
                display_text += f" ({prof})"
            self.sig_consulted_display.setText(display_text)
        else:
            self.sig_consulted_display.setText("Consulted professional: (set in Details panel)")

        # Update RC display
        rc = self.rc_name.text()
        if rc:
            self.sig_rc_display.setText(f"RC: {rc}")
        else:
            self.sig_rc_display.setText("RC: (set in Details panel)")

    def _update_clinical_card(self):
        text = self._generate_reasons_text()
        self.cards["clinical"].set_content_text(text if text else "")

    def _update_informal_card(self):
        text = self._generate_informal_text()
        self.cards["informal"].set_content_text(text if text else "")

    def _update_signature_card(self):
        parts = []
        # Show consulted person from clinician popup
        if self.consulted_name.text():
            parts.append(self.consulted_name.text())
            parts.append(self.prof_sig_date.date().toString('dd MMM yyyy'))
        # RC signature date
        parts.append(self.rc_sig_date.date().toString('dd MMM yyyy'))
        self.cards["signatures"].set_content_text("\n".join(parts))

    # ----------------------------------------------------------------
    # TEXT GENERATION
    # ----------------------------------------------------------------
    def _generate_reasons_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        paragraphs = []

        # Para 1: Demographics + Diagnosis
        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        if self.dx_primary.currentText().strip():
            diagnoses.append(self.dx_primary.currentText().strip())
        if self.dx_secondary.currentText().strip():
            diagnoses.append(self.dx_secondary.currentText().strip())

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for the patient to receive medical treatment in hospital.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for the patient to receive medical treatment in hospital.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for the patient to receive medical treatment in hospital.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        # Para 2: Necessity
        para2_parts = []
        necessity_items = []
        if self.nec_health.isChecked():
            necessity_items.append(f"{p['pos_l']} own health")
        if self.nec_safety.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append("own safety")
        if self.nec_safety.isChecked() and self.nec_others.isChecked():
            necessity_items.append("the protection of others")

        if necessity_items:
            joined = ", ".join(necessity_items[:-1]) + f" and {necessity_items[-1]}" if len(necessity_items) > 1 else necessity_items[0]
            para2_parts.append(f"The Mental Health Act is necessary for {joined}.")

        if self.nec_health.isChecked():
            mh_items = []
            if self.mental_health_cb.isChecked():
                if self.poor_compliance_cb.isChecked():
                    mh_items.append("poor compliance with treatment")
                if self.limited_insight_cb.isChecked():
                    mh_items.append(f"limited insight into {p['pos_l']} illness")
            if self.physical_health_cb.isChecked():
                ph_details = self.physical_health_details.text().strip()
                if ph_details:
                    mh_items.append(ph_details)
            if mh_items:
                para2_parts.append(f"Regarding {p['pos_l']} health, I would be concerned about {' and '.join(mh_items)} resulting in a deterioration in mental state if informal.")

        if self.nec_safety.isChecked() and self.self_harm_cb.isChecked():
            self_hist = []
            self_curr = []
            reflexive = "himself" if self.gender_male.isChecked() else ("herself" if self.gender_female.isChecked() else "themselves")
            if self.self_hist_neglect.isChecked(): self_hist.append("self neglect")
            if self.self_hist_risky.isChecked(): self_hist.append(f"placing of {reflexive} in risky situations")
            if self.self_hist_harm.isChecked(): self_hist.append("self harm")
            if self.self_curr_neglect.isChecked(): self_curr.append("self neglect")
            if self.self_curr_risky.isChecked(): self_curr.append(f"placing of {reflexive} in risky situations")
            if self.self_curr_harm.isChecked(): self_curr.append("self harm")

            if self_hist or self_curr:
                safety_text = f"With respect to {p['pos_l']} own safety"
                if self_hist:
                    hist_joined = ", ".join(self_hist[:-1]) + ", and of " + self_hist[-1] if len(self_hist) > 1 else self_hist[0]
                    safety_text += f", historically, when unwell and non-compliant there has been a risk of {hist_joined}"
                if self_curr:
                    if self_hist:
                        safety_text += f". This risk continues currently despite treatment"
                    else:
                        curr_joined = ", ".join(self_curr[:-1]) + ", and of " + self_curr[-1] if len(self_curr) > 1 else self_curr[0]
                        safety_text += f", currently there is a risk of {curr_joined}"
                safety_text += "."
                para2_parts.append(safety_text)

        if self.nec_safety.isChecked() and self.nec_others.isChecked():
            others_hist = []
            others_curr = []
            if self.others_hist_violence.isChecked(): others_hist.append("violence")
            if self.others_hist_verbal.isChecked(): others_hist.append("verbal aggression")
            if self.others_hist_sexual.isChecked(): others_hist.append("sexual violence")
            if self.others_hist_stalking.isChecked(): others_hist.append("stalking")
            if self.others_hist_arson.isChecked(): others_hist.append("arson")
            if self.others_curr_violence.isChecked(): others_curr.append("violence")
            if self.others_curr_verbal.isChecked(): others_curr.append("verbal aggression")
            if self.others_curr_sexual.isChecked(): others_curr.append("sexual violence")
            if self.others_curr_stalking.isChecked(): others_curr.append("stalking")
            if self.others_curr_arson.isChecked(): others_curr.append("arson")

            if others_hist or others_curr:
                both = [item for item in others_hist if item in others_curr]
                hist_only = [item for item in others_hist if item not in both]
                curr_only = [item for item in others_curr if item not in both]

                others_text = "Regarding risk to others, if not under the Mental Health Act and not compliant, I would be concerned about the risk of "
                all_items = []
                if both:
                    all_items.append(f"{', '.join(both)}, which is both historical and current")
                if hist_only:
                    all_items.append(f"{', '.join(hist_only)} (historical)")
                if curr_only:
                    all_items.append(f"{', '.join(curr_only)} (current)")
                if len(all_items) > 1:
                    others_text += ", and ".join([", ".join(all_items[:-1]), all_items[-1]]) + "."
                else:
                    others_text += all_items[0] + "."
                para2_parts.append(others_text)

        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        return "\n\n".join(paragraphs)

    def _generate_informal_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        parts = []
        if self.tried_failed_cb.isChecked():
            parts.append("Previous attempts at informal treatment have not been successful and I would be concerned about this recurring.")
        if self.insight_cb.isChecked():
            parts.append(f"{p['pos']} lack of insight is a significant concern and should {p['subj_l']} be discharged, I believe this would significantly impair {p['pos_l']} compliance.")
        if self.compliance_cb.isChecked():
            parts.append(f"Compliance with treatment has been a significant issue and I do not believe {p['subj_l']} would comply if informal.")
        if self.supervision_cb.isChecked():
            parts.append(f"I believe {name_display.lower() if name_display != 'The patient' else 'the patient'} needs the supervision afforded by the Mental Health Act.")

        if parts:
            return "Such treatment cannot be provided unless the patient continues to be detained under the Act. " + " ".join(parts)
        return ""

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.patient_name.clear()
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.hospital.clear()
            self.exam_date.setDate(QDate.currentDate())
            self.expiry_date.setDate(QDate.currentDate())
            self.rc_name.clear()
            self.rc_profession.setText("Consultant Psychiatrist")
            self.rc_address.clear()
            self.rc_email.clear()
            self.consulted_name.clear()
            self.consulted_profession.setCurrentIndex(0)
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.degree_details.clear()
            self.nec_health.setChecked(False)
            self.nec_safety.setChecked(False)
            self.physical_health_details.clear()
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.prof_sig_date.setDate(QDate.currentDate())
            self.rc_sig_date.setDate(QDate.currentDate())
            self.furnish_internal.setChecked(True)
            # Clear all cards
            for card in self.cards.values():
                card.set_content_text("")
            # Restore my details fields
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form H5",
            f"Form_H5_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_H5_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form H5 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color (#918C0D) and cream highlight (#FFFED5)
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)
            CREAM_FILL = 'FFFED5'

            # Pre-process: clean ALL paragraphs - remove permission markers, convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), CREAM_FILL)
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), CREAM_FILL)

            def set_entry_box(para, content: str):
                """Set entry box with gold brackets."""
                if not content or not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Opening bracket
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def format_option_first(para, content: str, strike: bool = False):
                """First option - opening bracket only."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = '['
                    ob = para.runs[0]
                else:
                    ob = para.add_run('[')
                ob.font.name = 'Arial'
                ob.font.size = Pt(12)
                ob.font.bold = True
                ob.font.color.rgb = BRACKET_COLOR
                rPr_ob = ob._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                if strike:
                    ob.font.strike = True
                ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True

            def format_option_middle(para, content: str, strike: bool = False):
                """Middle option - no brackets, just cream highlight."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True

            def format_option_last(para, content: str, strike: bool = False):
                """Last option - closing bracket only."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True
                cb = para.add_run(']')
                cb.font.name = 'Arial'
                cb.font.size = Pt(12)
                cb.font.bold = True
                cb.font.color.rgb = BRACKET_COLOR
                rPr_cb = cb._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)
                if strike:
                    cb.font.strike = True

            def format_cream_para(para, content: str):
                """Cream highlight paragraph."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)

            def format_closing_bracket_para(para):
                """Paragraph with cream and closing bracket."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = '                                                                                                    '
                    ct = para.runs[0]
                else:
                    ct = para.add_run('                                                                                                    ')
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                cb = para.add_run(']')
                cb.font.name = 'Arial'
                cb.font.size = Pt(12)
                cb.font.bold = True
                cb.font.color.rgb = BRACKET_COLOR
                rPr_cb = cb._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def format_sig_line(para, label1, label2, value2=""):
                """Format signature line: Label1 [ ] Label2 [ value2 ]"""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Label1
                if para.runs:
                    para.runs[0].text = label1
                    l1 = para.runs[0]
                else:
                    l1 = para.add_run(label1)
                l1.font.name = 'Arial'
                l1.font.size = Pt(12)
                # First placeholder
                ob1 = para.add_run('[')
                ob1.font.bold = True
                ob1.font.color.rgb = BRACKET_COLOR
                rPr_ob1 = ob1._element.get_or_add_rPr()
                shd_ob1 = OxmlElement('w:shd')
                shd_ob1.set(qn('w:val'), 'clear')
                shd_ob1.set(qn('w:color'), 'auto')
                shd_ob1.set(qn('w:fill'), CREAM_FILL)
                rPr_ob1.append(shd_ob1)
                c1 = para.add_run('                              ')
                rPr_c1 = c1._element.get_or_add_rPr()
                shd_c1 = OxmlElement('w:shd')
                shd_c1.set(qn('w:val'), 'clear')
                shd_c1.set(qn('w:color'), 'auto')
                shd_c1.set(qn('w:fill'), CREAM_FILL)
                rPr_c1.append(shd_c1)
                cb1 = para.add_run(']')
                cb1.font.bold = True
                cb1.font.color.rgb = BRACKET_COLOR
                rPr_cb1 = cb1._element.get_or_add_rPr()
                shd_cb1 = OxmlElement('w:shd')
                shd_cb1.set(qn('w:val'), 'clear')
                shd_cb1.set(qn('w:color'), 'auto')
                shd_cb1.set(qn('w:fill'), CREAM_FILL)
                rPr_cb1.append(shd_cb1)
                # Label2
                l2 = para.add_run(' ' + label2)
                l2.font.name = 'Arial'
                l2.font.size = Pt(12)
                # Second placeholder
                ob2 = para.add_run('[')
                ob2.font.bold = True
                ob2.font.color.rgb = BRACKET_COLOR
                rPr_ob2 = ob2._element.get_or_add_rPr()
                shd_ob2 = OxmlElement('w:shd')
                shd_ob2.set(qn('w:val'), 'clear')
                shd_ob2.set(qn('w:color'), 'auto')
                shd_ob2.set(qn('w:fill'), CREAM_FILL)
                rPr_ob2.append(shd_ob2)
                content2 = value2 if value2 else '                              '
                c2 = para.add_run(content2)
                rPr_c2 = c2._element.get_or_add_rPr()
                shd_c2 = OxmlElement('w:shd')
                shd_c2.set(qn('w:val'), 'clear')
                shd_c2.set(qn('w:color'), 'auto')
                shd_c2.set(qn('w:fill'), CREAM_FILL)
                rPr_c2.append(shd_c2)
                cb2 = para.add_run(']')
                cb2.font.bold = True
                cb2.font.color.rgb = BRACKET_COLOR
                rPr_cb2 = cb2._element.get_or_add_rPr()
                shd_cb2 = OxmlElement('w:shd')
                shd_cb2.set(qn('w:val'), 'clear')
                shd_cb2.set(qn('w:color'), 'auto')
                shd_cb2.set(qn('w:fill'), CREAM_FILL)
                rPr_cb2.append(shd_cb2)

            paragraphs = doc.paragraphs

            # Hospital (para 5)
            hospital_text = self.hospital.text().strip()
            set_entry_box(paragraphs[5], hospital_text)

            # Patient name (para 7)
            patient_text = self.patient_name.text().strip()
            set_entry_box(paragraphs[7], patient_text)

            # Exam date (para 9)
            exam_date = self.exam_date.date().toString("dd/MM/yyyy")
            set_entry_box(paragraphs[9], exam_date)

            # Expiry date (para 11)
            expiry_date = self.expiry_date.date().toString("dd/MM/yyyy")
            set_entry_box(paragraphs[11], expiry_date)

            # Consulted (para 13)
            consult_text = self.consulted_name.text()
            if self.consulted_profession.currentIndex() > 0:
                consult_text += ", " + self.consulted_profession.currentText()
            set_entry_box(paragraphs[13], consult_text)

            # (i), (ii), (iii) Necessity options - Para 19, 20, 21
            any_nec_selected = self.nec_health.isChecked() or self.nec_safety.isChecked() or self.nec_others.isChecked()

            strike_health = any_nec_selected and not self.nec_health.isChecked()
            strike_safety = any_nec_selected and not self.nec_safety.isChecked()
            strike_others = any_nec_selected and not self.nec_others.isChecked()

            format_option_first(paragraphs[19], "for the patient's own health", strike=strike_health)
            format_option_middle(paragraphs[20], "for the patient's own safety", strike=strike_safety)
            format_option_last(paragraphs[21], "for the protection of other persons", strike=strike_others)

            # First reasons box - Para 25-27 (opening bracket on 25, closing on 27)
            reasons_text = self.cards["clinical"].get_content().strip()
            if not reasons_text:
                reasons_text = '                                                                                                    '

            # Para 25 - opening bracket + content
            for run in paragraphs[25].runs:
                run.text = ""
            while len(paragraphs[25].runs) > 1:
                paragraphs[25]._element.remove(paragraphs[25].runs[-1]._element)
            pPr25 = paragraphs[25]._element.get_or_add_pPr()
            for old_shd in pPr25.findall(qn('w:shd')):
                pPr25.remove(old_shd)
            if paragraphs[25].runs:
                paragraphs[25].runs[0].text = '['
                ob25 = paragraphs[25].runs[0]
            else:
                ob25 = paragraphs[25].add_run('[')
            ob25.font.name = 'Arial'
            ob25.font.size = Pt(12)
            ob25.font.bold = True
            ob25.font.color.rgb = BRACKET_COLOR
            rPr_ob25 = ob25._element.get_or_add_rPr()
            shd_ob25 = OxmlElement('w:shd')
            shd_ob25.set(qn('w:val'), 'clear')
            shd_ob25.set(qn('w:color'), 'auto')
            shd_ob25.set(qn('w:fill'), CREAM_FILL)
            rPr_ob25.append(shd_ob25)
            ct25 = paragraphs[25].add_run(reasons_text)
            ct25.font.name = 'Arial'
            ct25.font.size = Pt(12)
            rPr_ct25 = ct25._element.get_or_add_rPr()
            shd_ct25 = OxmlElement('w:shd')
            shd_ct25.set(qn('w:val'), 'clear')
            shd_ct25.set(qn('w:color'), 'auto')
            shd_ct25.set(qn('w:fill'), CREAM_FILL)
            rPr_ct25.append(shd_ct25)

            # Para 26 - cream continuation
            format_cream_para(paragraphs[26], '                                                                                                    ')

            # Para 27 - cream + closing bracket
            format_closing_bracket_para(paragraphs[27])

            # Second reasons box - Para 30-32 (opening bracket on 30, closing on 32)
            informal_text = self.cards["informal"].get_content().strip()
            if not informal_text:
                informal_text = '                                                                                                    '

            # Para 30 - opening bracket + content
            for run in paragraphs[30].runs:
                run.text = ""
            while len(paragraphs[30].runs) > 1:
                paragraphs[30]._element.remove(paragraphs[30].runs[-1]._element)
            pPr30 = paragraphs[30]._element.get_or_add_pPr()
            for old_shd in pPr30.findall(qn('w:shd')):
                pPr30.remove(old_shd)
            if paragraphs[30].runs:
                paragraphs[30].runs[0].text = '['
                ob30 = paragraphs[30].runs[0]
            else:
                ob30 = paragraphs[30].add_run('[')
            ob30.font.name = 'Arial'
            ob30.font.size = Pt(12)
            ob30.font.bold = True
            ob30.font.color.rgb = BRACKET_COLOR
            rPr_ob30 = ob30._element.get_or_add_rPr()
            shd_ob30 = OxmlElement('w:shd')
            shd_ob30.set(qn('w:val'), 'clear')
            shd_ob30.set(qn('w:color'), 'auto')
            shd_ob30.set(qn('w:fill'), CREAM_FILL)
            rPr_ob30.append(shd_ob30)
            ct30 = paragraphs[30].add_run(informal_text)
            ct30.font.name = 'Arial'
            ct30.font.size = Pt(12)
            rPr_ct30 = ct30._element.get_or_add_rPr()
            shd_ct30 = OxmlElement('w:shd')
            shd_ct30.set(qn('w:val'), 'clear')
            shd_ct30.set(qn('w:color'), 'auto')
            shd_ct30.set(qn('w:fill'), CREAM_FILL)
            rPr_ct30.append(shd_ct30)

            # Para 31 - cream continuation
            format_cream_para(paragraphs[31], '                                                                                                    ')

            # Para 32 - cream + closing bracket
            format_closing_bracket_para(paragraphs[32])

            # PART 1 Signature lines - Para 36, 37
            rc_name = self.rc_name.text().strip()
            rc_profession = self.rc_profession.text().strip()
            rc_date = self.rc_sig_date.date().toString("dd/MM/yyyy")
            format_sig_line(paragraphs[36], "Signed", "PRINT NAME", rc_name)
            # Para 37: Profession [rc_profession] Date [rc_date]
            for run in paragraphs[37].runs:
                run.text = ""
            while len(paragraphs[37].runs) > 1:
                paragraphs[37]._element.remove(paragraphs[37].runs[-1]._element)
            pPr37 = paragraphs[37]._element.get_or_add_pPr()
            for old_shd in pPr37.findall(qn('w:shd')):
                pPr37.remove(old_shd)
            if paragraphs[37].runs:
                paragraphs[37].runs[0].text = 'Profession'
                l37 = paragraphs[37].runs[0]
            else:
                l37 = paragraphs[37].add_run('Profession')
            l37.font.name = 'Arial'
            l37.font.size = Pt(12)
            ob37a = paragraphs[37].add_run('[')
            ob37a.font.bold = True
            ob37a.font.color.rgb = BRACKET_COLOR
            rPr_ob37a = ob37a._element.get_or_add_rPr()
            shd_ob37a = OxmlElement('w:shd')
            shd_ob37a.set(qn('w:val'), 'clear')
            shd_ob37a.set(qn('w:color'), 'auto')
            shd_ob37a.set(qn('w:fill'), CREAM_FILL)
            rPr_ob37a.append(shd_ob37a)
            prof_content = rc_profession if rc_profession else '                              '
            c37a = paragraphs[37].add_run(prof_content)
            rPr_c37a = c37a._element.get_or_add_rPr()
            shd_c37a = OxmlElement('w:shd')
            shd_c37a.set(qn('w:val'), 'clear')
            shd_c37a.set(qn('w:color'), 'auto')
            shd_c37a.set(qn('w:fill'), CREAM_FILL)
            rPr_c37a.append(shd_c37a)
            cb37a = paragraphs[37].add_run(']')
            cb37a.font.bold = True
            cb37a.font.color.rgb = BRACKET_COLOR
            rPr_cb37a = cb37a._element.get_or_add_rPr()
            shd_cb37a = OxmlElement('w:shd')
            shd_cb37a.set(qn('w:val'), 'clear')
            shd_cb37a.set(qn('w:color'), 'auto')
            shd_cb37a.set(qn('w:fill'), CREAM_FILL)
            rPr_cb37a.append(shd_cb37a)
            dl37 = paragraphs[37].add_run(' Date')
            dl37.font.name = 'Arial'
            dl37.font.size = Pt(12)
            ob37b = paragraphs[37].add_run('[')
            ob37b.font.bold = True
            ob37b.font.color.rgb = BRACKET_COLOR
            rPr_ob37b = ob37b._element.get_or_add_rPr()
            shd_ob37b = OxmlElement('w:shd')
            shd_ob37b.set(qn('w:val'), 'clear')
            shd_ob37b.set(qn('w:color'), 'auto')
            shd_ob37b.set(qn('w:fill'), CREAM_FILL)
            rPr_ob37b.append(shd_ob37b)
            c37b = paragraphs[37].add_run(rc_date if rc_date else '                              ')
            rPr_c37b = c37b._element.get_or_add_rPr()
            shd_c37b = OxmlElement('w:shd')
            shd_c37b.set(qn('w:val'), 'clear')
            shd_c37b.set(qn('w:color'), 'auto')
            shd_c37b.set(qn('w:fill'), CREAM_FILL)
            rPr_c37b.append(shd_c37b)
            cb37b = paragraphs[37].add_run(']')
            cb37b.font.bold = True
            cb37b.font.color.rgb = BRACKET_COLOR
            rPr_cb37b = cb37b._element.get_or_add_rPr()
            shd_cb37b = OxmlElement('w:shd')
            shd_cb37b.set(qn('w:val'), 'clear')
            shd_cb37b.set(qn('w:color'), 'auto')
            shd_cb37b.set(qn('w:fill'), CREAM_FILL)
            rPr_cb37b.append(shd_cb37b)

            # PART 2 Signature lines - Para 41, 42
            consulted_name = self.consulted_name.text().strip()
            consulted_prof = self.consulted_profession.currentText() if self.consulted_profession.currentIndex() > 0 else ""
            prof_date = self.prof_sig_date.date().toString("dd/MM/yyyy")
            format_sig_line(paragraphs[41], "Signed", "PRINT NAME", consulted_name)
            # Para 42: Profession [consulted_prof] Date [prof_date]
            for run in paragraphs[42].runs:
                run.text = ""
            while len(paragraphs[42].runs) > 1:
                paragraphs[42]._element.remove(paragraphs[42].runs[-1]._element)
            pPr42 = paragraphs[42]._element.get_or_add_pPr()
            for old_shd in pPr42.findall(qn('w:shd')):
                pPr42.remove(old_shd)
            if paragraphs[42].runs:
                paragraphs[42].runs[0].text = 'Profession'
                l42 = paragraphs[42].runs[0]
            else:
                l42 = paragraphs[42].add_run('Profession')
            l42.font.name = 'Arial'
            l42.font.size = Pt(12)
            ob42a = paragraphs[42].add_run('[')
            ob42a.font.bold = True
            ob42a.font.color.rgb = BRACKET_COLOR
            rPr_ob42a = ob42a._element.get_or_add_rPr()
            shd_ob42a = OxmlElement('w:shd')
            shd_ob42a.set(qn('w:val'), 'clear')
            shd_ob42a.set(qn('w:color'), 'auto')
            shd_ob42a.set(qn('w:fill'), CREAM_FILL)
            rPr_ob42a.append(shd_ob42a)
            prof2_content = consulted_prof if consulted_prof else '                              '
            c42a = paragraphs[42].add_run(prof2_content)
            rPr_c42a = c42a._element.get_or_add_rPr()
            shd_c42a = OxmlElement('w:shd')
            shd_c42a.set(qn('w:val'), 'clear')
            shd_c42a.set(qn('w:color'), 'auto')
            shd_c42a.set(qn('w:fill'), CREAM_FILL)
            rPr_c42a.append(shd_c42a)
            cb42a = paragraphs[42].add_run(']')
            cb42a.font.bold = True
            cb42a.font.color.rgb = BRACKET_COLOR
            rPr_cb42a = cb42a._element.get_or_add_rPr()
            shd_cb42a = OxmlElement('w:shd')
            shd_cb42a.set(qn('w:val'), 'clear')
            shd_cb42a.set(qn('w:color'), 'auto')
            shd_cb42a.set(qn('w:fill'), CREAM_FILL)
            rPr_cb42a.append(shd_cb42a)
            dl42 = paragraphs[42].add_run(' Date')
            dl42.font.name = 'Arial'
            dl42.font.size = Pt(12)
            ob42b = paragraphs[42].add_run('[')
            ob42b.font.bold = True
            ob42b.font.color.rgb = BRACKET_COLOR
            rPr_ob42b = ob42b._element.get_or_add_rPr()
            shd_ob42b = OxmlElement('w:shd')
            shd_ob42b.set(qn('w:val'), 'clear')
            shd_ob42b.set(qn('w:color'), 'auto')
            shd_ob42b.set(qn('w:fill'), CREAM_FILL)
            rPr_ob42b.append(shd_ob42b)
            c42b = paragraphs[42].add_run(prof_date if prof_date else '                              ')
            rPr_c42b = c42b._element.get_or_add_rPr()
            shd_c42b = OxmlElement('w:shd')
            shd_c42b.set(qn('w:val'), 'clear')
            shd_c42b.set(qn('w:color'), 'auto')
            shd_c42b.set(qn('w:fill'), CREAM_FILL)
            rPr_c42b.append(shd_c42b)
            cb42b = paragraphs[42].add_run(']')
            cb42b.font.bold = True
            cb42b.font.color.rgb = BRACKET_COLOR
            rPr_cb42b = cb42b._element.get_or_add_rPr()
            shd_cb42b = OxmlElement('w:shd')
            shd_cb42b.set(qn('w:val'), 'clear')
            shd_cb42b.set(qn('w:color'), 'auto')
            shd_cb42b.set(qn('w:fill'), CREAM_FILL)
            rPr_cb42b.append(shd_cb42b)

            # PART 3 Furnishing options - Para 46, 47, 48
            any_furnish_selected = self.furnish_internal.isChecked() or self.furnish_electronic.isChecked() or self.furnish_other.isChecked()

            strike_internal = any_furnish_selected and not self.furnish_internal.isChecked()
            strike_electronic = any_furnish_selected and not self.furnish_electronic.isChecked()
            strike_other = any_furnish_selected and not self.furnish_other.isChecked()

            format_option_first(paragraphs[46], "today consigning it to the hospital managers' internal mail system.", strike=strike_internal)
            format_option_middle(paragraphs[47], "today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.", strike=strike_electronic)
            format_option_last(paragraphs[48], "sending or delivering it without using the hospital managers' internal mail system.", strike=strike_other)

            # PART 3 Signed line - Para 49 (single placeholder only)
            for run in paragraphs[49].runs:
                run.text = ""
            while len(paragraphs[49].runs) > 1:
                paragraphs[49]._element.remove(paragraphs[49].runs[-1]._element)
            pPr49 = paragraphs[49]._element.get_or_add_pPr()
            for old_shd in pPr49.findall(qn('w:shd')):
                pPr49.remove(old_shd)
            if paragraphs[49].runs:
                paragraphs[49].runs[0].text = 'Signed'
                sl49 = paragraphs[49].runs[0]
            else:
                sl49 = paragraphs[49].add_run('Signed')
            sl49.font.name = 'Arial'
            sl49.font.size = Pt(12)
            ob49 = paragraphs[49].add_run('[')
            ob49.font.bold = True
            ob49.font.color.rgb = BRACKET_COLOR
            rPr_ob49 = ob49._element.get_or_add_rPr()
            shd_ob49 = OxmlElement('w:shd')
            shd_ob49.set(qn('w:val'), 'clear')
            shd_ob49.set(qn('w:color'), 'auto')
            shd_ob49.set(qn('w:fill'), CREAM_FILL)
            rPr_ob49.append(shd_ob49)
            c49 = paragraphs[49].add_run('                                                      ')
            rPr_c49 = c49._element.get_or_add_rPr()
            shd_c49 = OxmlElement('w:shd')
            shd_c49.set(qn('w:val'), 'clear')
            shd_c49.set(qn('w:color'), 'auto')
            shd_c49.set(qn('w:fill'), CREAM_FILL)
            rPr_c49.append(shd_c49)
            cb49 = paragraphs[49].add_run(']')
            cb49.font.bold = True
            cb49.font.color.rgb = BRACKET_COLOR
            rPr_cb49 = cb49._element.get_or_add_rPr()
            shd_cb49 = OxmlElement('w:shd')
            shd_cb49.set(qn('w:val'), 'clear')
            shd_cb49.set(qn('w:color'), 'auto')
            shd_cb49.set(qn('w:fill'), CREAM_FILL)
            rPr_cb49.append(shd_cb49)

            # PART 3 PRINT NAME/Date line - Para 50
            format_sig_line(paragraphs[50], "PRINT NAME", "Date", rc_date)

            # PART 4 Receipt options - Para 54, 55, 56 (no strikethrough in Part 4)
            format_option_first(paragraphs[54], "furnished to the hospital managers through their internal mail system.", strike=False)
            format_option_middle(paragraphs[55], "furnished to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.", strike=False)
            format_option_middle(paragraphs[56], "received by me on behalf of the hospital managers on [date].", strike=False)

            # Para 57 - placeholder with closing bracket
            format_closing_bracket_para(paragraphs[57])

            # PART 4 Signed line - Para 58
            for run in paragraphs[58].runs:
                run.text = ""
            while len(paragraphs[58].runs) > 1:
                paragraphs[58]._element.remove(paragraphs[58].runs[-1]._element)
            pPr58 = paragraphs[58]._element.get_or_add_pPr()
            for old_shd in pPr58.findall(qn('w:shd')):
                pPr58.remove(old_shd)
            if paragraphs[58].runs:
                paragraphs[58].runs[0].text = 'Signed'
                sl58 = paragraphs[58].runs[0]
            else:
                sl58 = paragraphs[58].add_run('Signed')
            sl58.font.name = 'Arial'
            sl58.font.size = Pt(12)
            ob58 = paragraphs[58].add_run('[')
            ob58.font.bold = True
            ob58.font.color.rgb = BRACKET_COLOR
            rPr_ob58 = ob58._element.get_or_add_rPr()
            shd_ob58 = OxmlElement('w:shd')
            shd_ob58.set(qn('w:val'), 'clear')
            shd_ob58.set(qn('w:color'), 'auto')
            shd_ob58.set(qn('w:fill'), CREAM_FILL)
            rPr_ob58.append(shd_ob58)
            c58 = paragraphs[58].add_run('                                        ')
            rPr_c58 = c58._element.get_or_add_rPr()
            shd_c58 = OxmlElement('w:shd')
            shd_c58.set(qn('w:val'), 'clear')
            shd_c58.set(qn('w:color'), 'auto')
            shd_c58.set(qn('w:fill'), CREAM_FILL)
            rPr_c58.append(shd_c58)
            cb58 = paragraphs[58].add_run(']')
            cb58.font.bold = True
            cb58.font.color.rgb = BRACKET_COLOR
            rPr_cb58 = cb58._element.get_or_add_rPr()
            shd_cb58 = OxmlElement('w:shd')
            shd_cb58.set(qn('w:val'), 'clear')
            shd_cb58.set(qn('w:color'), 'auto')
            shd_cb58.set(qn('w:fill'), CREAM_FILL)
            rPr_cb58.append(shd_cb58)
            suf58 = paragraphs[58].add_run(' on behalf of the hospital managers')
            suf58.font.name = 'Arial'
            suf58.font.size = Pt(12)

            # PART 4 PRINT NAME/Date line - Para 59
            format_sig_line(paragraphs[59], "PRINT NAME", "Date", "")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form H5 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        gender = "neutral"
        if self.gender_male.isChecked():
            gender = "male"
        elif self.gender_female.isChecked():
            gender = "female"

        return {
            "patient_name": self.patient_name.text(),
            "age": self.age_spin.value(),
            "gender": gender,
            "ethnicity": self.ethnicity_combo.currentText(),
            "hospital": self.hospital.text(),
            "exam_date": self.exam_date.date().toString("yyyy-MM-dd"),
            "expiry_date": self.expiry_date.date().toString("yyyy-MM-dd"),
            "rc_name": self.rc_name.text(),
            "rc_profession": self.rc_profession.text(),
            "rc_address": self.rc_address.text(),
            "rc_email": self.rc_email.text(),
            "dx_primary": self.dx_primary.currentText(),
            "dx_secondary": self.dx_secondary.currentText(),
            "consulted_name": self.consulted_name.text(),
            "consulted_profession": self.consulted_profession.currentText(),
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "nec_health": self.nec_health.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "nec_safety": self.nec_safety.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_hist_neglect": self.self_hist_neglect.isChecked(),
            "self_hist_risky": self.self_hist_risky.isChecked(),
            "self_hist_harm": self.self_hist_harm.isChecked(),
            "self_curr_neglect": self.self_curr_neglect.isChecked(),
            "self_curr_risky": self.self_curr_risky.isChecked(),
            "self_curr_harm": self.self_curr_harm.isChecked(),
            "nec_others": self.nec_others.isChecked(),
            "others_hist_violence": self.others_hist_violence.isChecked(),
            "others_hist_verbal": self.others_hist_verbal.isChecked(),
            "others_hist_sexual": self.others_hist_sexual.isChecked(),
            "others_hist_stalking": self.others_hist_stalking.isChecked(),
            "others_hist_arson": self.others_hist_arson.isChecked(),
            "others_curr_violence": self.others_curr_violence.isChecked(),
            "others_curr_verbal": self.others_curr_verbal.isChecked(),
            "others_curr_sexual": self.others_curr_sexual.isChecked(),
            "others_curr_stalking": self.others_curr_stalking.isChecked(),
            "others_curr_arson": self.others_curr_arson.isChecked(),
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "prof_sig_date": self.prof_sig_date.date().toString("yyyy-MM-dd"),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "furnish_method": "internal" if self.furnish_internal.isChecked() else ("electronic" if self.furnish_electronic.isChecked() else "other"),
            # Card contents
            "card_details": self.cards["details"].get_content(),
            "card_clinical": self.cards["clinical"].get_content(),
            "card_informal": self.cards["informal"].get_content(),
            "card_signatures": self.cards["signatures"].get_content(),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.patient_name.setText(state.get("patient_name", ""))
        self.age_spin.setValue(state.get("age", 0))
        g = state.get("gender", "neutral")
        if g == "male":
            self.gender_male.setChecked(True)
        elif g == "female":
            self.gender_female.setChecked(True)
        else:
            self.gender_other.setChecked(True)
        idx = self.ethnicity_combo.findText(state.get("ethnicity", "Not specified"))
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.hospital.setText(state.get("hospital", ""))
        if state.get("exam_date"):
            self.exam_date.setDate(QDate.fromString(state["exam_date"], "yyyy-MM-dd"))
        if state.get("expiry_date"):
            self.expiry_date.setDate(QDate.fromString(state["expiry_date"], "yyyy-MM-dd"))
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_profession.setText(state.get("rc_profession", "Consultant Psychiatrist"))
        self.rc_address.setText(state.get("rc_address", ""))
        self.rc_email.setText(state.get("rc_email", ""))
        self.consulted_name.setText(state.get("consulted_name", ""))
        prof_idx = self.consulted_profession.findText(state.get("consulted_profession", ""))
        if prof_idx >= 0:
            self.consulted_profession.setCurrentIndex(prof_idx)
        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)
        dx_secondary = state.get("dx_secondary", "")
        idx = self.dx_secondary.findText(dx_secondary)
        if idx >= 0:
            self.dx_secondary.setCurrentIndex(idx)
        else:
            self.dx_secondary.setCurrentText(dx_secondary)
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.nec_health.setChecked(state.get("nec_health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.nec_safety.setChecked(state.get("nec_safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_hist_neglect.setChecked(state.get("self_hist_neglect", False))
        self.self_hist_risky.setChecked(state.get("self_hist_risky", False))
        self.self_hist_harm.setChecked(state.get("self_hist_harm", False))
        self.self_curr_neglect.setChecked(state.get("self_curr_neglect", False))
        self.self_curr_risky.setChecked(state.get("self_curr_risky", False))
        self.self_curr_harm.setChecked(state.get("self_curr_harm", False))
        self.nec_others.setChecked(state.get("nec_others", False))
        self.others_hist_violence.setChecked(state.get("others_hist_violence", False))
        self.others_hist_verbal.setChecked(state.get("others_hist_verbal", False))
        self.others_hist_sexual.setChecked(state.get("others_hist_sexual", False))
        self.others_hist_stalking.setChecked(state.get("others_hist_stalking", False))
        self.others_hist_arson.setChecked(state.get("others_hist_arson", False))
        self.others_curr_violence.setChecked(state.get("others_curr_violence", False))
        self.others_curr_verbal.setChecked(state.get("others_curr_verbal", False))
        self.others_curr_sexual.setChecked(state.get("others_curr_sexual", False))
        self.others_curr_stalking.setChecked(state.get("others_curr_stalking", False))
        self.others_curr_arson.setChecked(state.get("others_curr_arson", False))
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        if state.get("prof_sig_date"):
            self.prof_sig_date.setDate(QDate.fromString(state["prof_sig_date"], "yyyy-MM-dd"))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        furnish = state.get("furnish_method", "internal")
        if furnish == "internal":
            self.furnish_internal.setChecked(True)
        elif furnish == "electronic":
            self.furnish_electronic.setChecked(True)
        else:
            self.furnish_other.setChecked(True)

        # Restore card contents
        if state.get("card_details"):
            self.cards["details"].set_content_text(state["card_details"])
        if state.get("card_clinical"):
            self.cards["clinical"].set_content_text(state["card_clinical"])
        if state.get("card_informal"):
            self.cards["informal"].set_content_text(state["card_informal"])
        if state.get("card_signatures"):
            self.cards["signatures"].set_content_text(state["card_signatures"])

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[H5Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[H5Form] Set gender: {gender}")
        # Set age if not already set (from age field or calculate from DOB)
        if hasattr(self, 'age_spin') and self.age_spin.value() == 0:
            age = patient_info.get("age")
            if not age and patient_info.get("dob"):
                from datetime import datetime
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if age and 0 < age < 120:
                self.age_spin.setValue(age)
                print(f"[H5Form] Set age: {age}")
        # Set ethnicity if not already selected
        if patient_info.get("ethnicity") and hasattr(self, 'ethnicity_combo'):
            current = self.ethnicity_combo.currentText()
            if current in ("Ethnicity", "Not specified", "Select..."):
                idx = self.ethnicity_combo.findText(patient_info["ethnicity"], Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    self.ethnicity_combo.setCurrentIndex(idx)
                    print(f"[H5Form] Set ethnicity: {patient_info['ethnicity']}")
