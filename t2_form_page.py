# ================================================================
#  T2 FORM PAGE — Certificate of Consent to Treatment
#  Mental Health Act 1983 - Form T2 Regulation 27(2)
#  Section 58(3)(a) — Certificate of consent to treatment
#  CARD/POPUP LAYOUT
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QPushButton, QFileDialog, QMessageBox, QToolButton,
    QComboBox, QRadioButton, QButtonGroup, QSplitter, QGroupBox,
    QStackedWidget, QSizePolicy
)
from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor
from background_history_popup import ResizableSection
from utils.resource_path import resource_path


# ================================================================
# NO-WHEEL WIDGETS (prevents scroll from changing value)
# ================================================================
class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# T2 CARD WIDGET
# ================================================================
class T2CardWidget(QFrame):
    """Clickable card with editable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("t2Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#t2Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#t2Card:hover {
                border-color: #059669;
                background: #ecfdf5;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        self.setMinimumHeight(80)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(6)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 19px; font-weight: 700; color: #059669;")
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 19px;
                color: #374151;
                background: transparent;
                border: none;
                padding: 0;
            }
        """)
        self.content.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        layout.addWidget(self.content, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=17)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_active(self, active: bool):
        """Set active state - shows persistent highlight."""
        if active:
            self.setStyleSheet("""
                QFrame#t2Card {
                    background: #ecfdf5;
                    border: 2px solid #059669;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#t2Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#t2Card:hover {
                    border-color: #059669;
                    background: #ecfdf5;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# T2 TOOLBAR
# ================================================================
class T2Toolbar(QWidget):
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            T2Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN T2 FORM PAGE - Card/Popup Layout
# ================================================================
class T2FormPage(QWidget):
    """Page for completing MHA Form T2 - Certificate of Consent to Treatment."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}
        self._regular_meds = []
        self._prn_meds = []

        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.clinician_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.clinician_email.setText(self._my_details["email"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(50)
        header.setStyleSheet("background: #059669; border-bottom: 1px solid #047857;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 5px;
                font-size: 18px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form T2 — Certificate of Consent to Treatment")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - cards on left, popup on right with splitter
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(12, 12, 12, 12)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #059669, stop:0.7 #059669, stop:1 transparent);
            }
        """)

        # Left: Cards column (scrollable)
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(16, 16, 16, 16)
        self.cards_layout.setSpacing(8)

        # Create cards
        self._create_clinician_card()
        self._create_patient_card()
        self._create_treatment_card()
        self._create_signature_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        self.main_splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 10px;
            }
        """)

        # Create popup panels
        self._create_clinician_popup()
        self._create_patient_popup()
        self._create_treatment_popup()
        self._create_signature_popup()

        # Initialize cards with default date values
        self._update_signature_card()

        self.main_splitter.addWidget(self.popup_stack)
        self.main_splitter.setSizes([280, 700])
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        # Show first popup by default
        self._on_card_clicked("clinician")

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"clinician": 0, "patient": 1, "treatment": 2, "signature": 3}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            for k, card in self.cards.items():
                card.set_active(k == key)

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_clinician_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 350
        card = T2CardWidget("Clinician Details", "clinician")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["clinician"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_patient_card(self):
        section = ResizableSection()
        section.set_content_height(180)
        section._min_height = 120
        section._max_height = 350
        card = T2CardWidget("Patient Details", "patient")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["patient"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_treatment_card(self):
        section = ResizableSection()
        section.set_content_height(200)
        section._min_height = 120
        section._max_height = 400
        card = T2CardWidget("Treatment", "treatment")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["treatment"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signature_card(self):
        section = ResizableSection()
        section.set_content_height(150)
        section._min_height = 100
        section._max_height = 300
        card = T2CardWidget("Signature", "signature")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signature"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_clinician_popup(self):
        """Popup for clinician details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Clinician Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #059669;")
        popup_layout.addWidget(header)

        # Scrollable form area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        form_widget = QWidget()
        form_widget.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        form_layout = QVBoxLayout(form_widget)
        form_layout.setContentsMargins(0, 0, 0, 0)
        form_layout.setSpacing(10)

        # Clinician type
        type_lbl = QLabel("I am:")
        type_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        form_layout.addWidget(type_lbl)

        self.clinician_type_group = QButtonGroup(self)
        self.approved_clinician_radio = QRadioButton("Approved clinician in charge of treatment")
        self.approved_clinician_radio.setChecked(True)
        self.approved_clinician_radio.setStyleSheet("font-size: 18px; color: #374151;")
        self.soad_radio = QRadioButton("SOAD (registered medical practitioner)")
        self.soad_radio.setStyleSheet("font-size: 18px; color: #374151;")
        self.clinician_type_group.addButton(self.approved_clinician_radio)
        self.clinician_type_group.addButton(self.soad_radio)
        form_layout.addWidget(self.approved_clinician_radio)
        form_layout.addWidget(self.soad_radio)

        # Name
        name_lbl = QLabel("Full Name:")
        name_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(name_lbl)
        self.clinician_name = self._create_line_edit("Clinician full name")
        form_layout.addWidget(self.clinician_name)

        # Address
        addr_lbl = QLabel("Address:")
        addr_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 4px;")
        form_layout.addWidget(addr_lbl)
        self.clinician_address = self._create_line_edit("Clinician address")
        form_layout.addWidget(self.clinician_address)

        # Email
        email_lbl = QLabel("Email:")
        email_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 4px;")
        form_layout.addWidget(email_lbl)
        self.clinician_email = self._create_line_edit("Clinician email")
        form_layout.addWidget(self.clinician_email)

        scroll.setWidget(form_widget)
        popup_layout.addWidget(scroll, 1)

        # Auto-sync to card
        self.approved_clinician_radio.toggled.connect(self._update_clinician_card)
        self.clinician_name.textChanged.connect(self._update_clinician_card)
        self.clinician_address.textChanged.connect(self._update_clinician_card)
        self.clinician_email.textChanged.connect(self._update_clinician_card)

        self.popup_stack.addWidget(popup)

    def _create_patient_popup(self):
        """Popup for patient details."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Patient Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #0891b2;")
        popup_layout.addWidget(header)

        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        # Patient name
        name_lbl = QLabel("Patient Full Name:")
        name_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        form_layout.addWidget(name_lbl)
        self.patient_name = self._create_line_edit("Patient full name")
        form_layout.addWidget(self.patient_name)

        # Patient address
        addr_lbl = QLabel("Patient Address:")
        addr_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151; margin-top: 8px;")
        form_layout.addWidget(addr_lbl)
        self.patient_address = self._create_line_edit("Patient address")
        form_layout.addWidget(self.patient_address)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        # Auto-sync to card
        self.patient_name.textChanged.connect(self._update_patient_card)
        self.patient_address.textChanged.connect(self._update_patient_card)

        self.popup_stack.addWidget(popup)

    def _create_treatment_popup(self):
        """Popup for treatment description with medication entries."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Treatment Description")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #7c3aed;")
        popup_layout.addWidget(header)

        info = QLabel("The patient is capable of understanding the nature, purpose and likely effects of the following treatment and has consented:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 17px; color: #6b7280; padding: 6px; background: #f3f4f6; border-radius: 4px;")
        popup_layout.addWidget(info)

        # Splitter: left = output text, right = medication inputs
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(6)
        splitter.setStyleSheet("QSplitter::handle { background: #d1d5db; } QSplitter::handle:hover { background: #6BAF8D; }")

        # Left side - output text area
        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 8, 0)

        output_lbl = QLabel("Treatment Summary:")
        output_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        left_layout.addWidget(output_lbl)

        self.treatment_desc = QTextEdit()
        self.treatment_desc.setPlaceholderText("Treatment summary will appear here...")
        self.treatment_desc.setReadOnly(True)
        self.treatment_desc.setStyleSheet("""
            QTextEdit {
                background: #1e1e1e;
                color: #d4d4d4;
                font-size: 17px;
                padding: 8px;
                border: 1px solid #3c3c3c;
                border-radius: 6px;
            }
        """)
        left_layout.addWidget(self.treatment_desc, 1)
        splitter.addWidget(left_widget)

        # Right side - medication inputs (scrollable)
        right_scroll = QScrollArea()
        right_scroll.setWidgetResizable(True)
        right_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        right_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        right_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        right_widget = QWidget()
        right_widget.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(8, 0, 0, 0)
        right_layout.setSpacing(12)

        # Regular Medications section
        regular_frame = QFrame()
        regular_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        regular_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 6px; }")
        regular_layout = QVBoxLayout(regular_frame)
        regular_layout.setContentsMargins(10, 8, 10, 8)
        regular_layout.setSpacing(6)

        regular_header = QLabel("Regular Medications")
        regular_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #166534;")
        regular_layout.addWidget(regular_header)

        self.regular_meds_container = QVBoxLayout()
        self.regular_meds_container.setSpacing(4)
        regular_layout.addLayout(self.regular_meds_container)

        add_regular_btn = QPushButton("+ Add Regular Med")
        add_regular_btn.setStyleSheet("""
            QPushButton {
                background: #dcfce7;
                color: #166534;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: #bbf7d0; }
        """)
        add_regular_btn.clicked.connect(lambda: self._add_medication_entry("regular"))
        regular_layout.addWidget(add_regular_btn)
        right_layout.addWidget(regular_frame)

        # PRN Medications section
        prn_frame = QFrame()
        prn_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        prn_frame.setStyleSheet("QFrame { background: #fef3c7; border: none; border-radius: 6px; }")
        prn_layout = QVBoxLayout(prn_frame)
        prn_layout.setContentsMargins(10, 8, 10, 8)
        prn_layout.setSpacing(6)

        prn_header = QLabel("PRN Medications")
        prn_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #92400e;")
        prn_layout.addWidget(prn_header)

        self.prn_meds_container = QVBoxLayout()
        self.prn_meds_container.setSpacing(4)
        prn_layout.addLayout(self.prn_meds_container)

        add_prn_btn = QPushButton("+ Add PRN Med")
        add_prn_btn.setStyleSheet("""
            QPushButton {
                background: #fef9c3;
                color: #92400e;
                border: none;
                padding: 6px 12px;
                border-radius: 4px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: #fde68a; }
        """)
        add_prn_btn.clicked.connect(lambda: self._add_medication_entry("prn"))
        prn_layout.addWidget(add_prn_btn)
        right_layout.addWidget(prn_frame)

        right_layout.addStretch()
        right_scroll.setWidget(right_widget)
        splitter.addWidget(right_scroll)
        splitter.setSizes([350, 350])

        popup_layout.addWidget(splitter, 1)

        self.popup_stack.addWidget(popup)

        # Add initial medication entries
        self._add_medication_entry("regular")
        self._add_medication_entry("prn")

    def _create_signature_popup(self):
        """Popup for signature."""
        popup = QWidget()
        popup_layout = QVBoxLayout(popup)
        popup_layout.setContentsMargins(16, 16, 16, 16)
        popup_layout.setSpacing(12)

        header = QLabel("Signature")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #dc2626;")
        popup_layout.addWidget(header)

        form_layout = QVBoxLayout()
        form_layout.setSpacing(10)

        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        form_layout.addWidget(sig_lbl)

        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(160)
        form_layout.addWidget(self.sig_date)

        popup_layout.addLayout(form_layout)
        popup_layout.addStretch()

        # Auto-sync to card
        self.sig_date.dateChanged.connect(self._update_signature_card)

        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # MEDICATION ENTRY
    # ----------------------------------------------------------------
    def _add_medication_entry(self, med_type: str):
        """Add a medication entry row (regular or prn)."""
        try:
            from CANONICAL_MEDS import MEDICATIONS
        except ImportError:
            MEDICATIONS = {}

        entry_widget = QFrame()
        entry_widget.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Fixed)
        entry_widget.setStyleSheet("""
            QFrame {
                background: white;
                border-radius: 4px;
                border: 1px solid #e5e7eb;
            }
        """)
        entry_layout = QHBoxLayout(entry_widget)
        entry_layout.setContentsMargins(6, 4, 6, 4)
        entry_layout.setSpacing(6)

        # Medication dropdown
        med_combo = NoWheelComboBox()
        med_combo.setEditable(True)
        med_combo.addItem("")
        med_combo.addItems(sorted(MEDICATIONS.keys()))
        med_combo.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Fixed)
        med_combo.setMinimumWidth(80)
        med_combo.setStyleSheet("""
            QComboBox {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 4px 6px;
                font-size: 17px;
            }
        """)
        entry_layout.addWidget(med_combo, 1)

        # BNF radio buttons
        bnf_group = QButtonGroup(entry_widget)
        bnf_radio = QRadioButton("BNF")
        bnf_radio.setChecked(True)
        bnf_radio.setStyleSheet("font-size: 16px;")
        above_bnf_radio = QRadioButton("Above")
        above_bnf_radio.setStyleSheet("font-size: 16px;")
        bnf_group.addButton(bnf_radio)
        bnf_group.addButton(above_bnf_radio)
        entry_layout.addWidget(bnf_radio)
        entry_layout.addWidget(above_bnf_radio)

        # Remove button
        remove_btn = QPushButton("×")
        remove_btn.setFixedSize(18, 18)
        remove_btn.setStyleSheet("""
            QPushButton {
                background: #ef4444;
                color: white;
                border: none;
                border-radius: 9px;
                font-size: 18px;
                font-weight: bold;
            }
            QPushButton:hover { background: #dc2626; }
        """)
        entry_layout.addWidget(remove_btn)

        # Store entry data
        entry_data = {
            "widget": entry_widget,
            "name": med_combo,
            "bnf_radio": bnf_radio,
            "above_bnf_radio": above_bnf_radio
        }

        if med_type == "regular":
            self._regular_meds.append(entry_data)
            self.regular_meds_container.addWidget(entry_widget)
        else:
            self._prn_meds.append(entry_data)
            self.prn_meds_container.addWidget(entry_widget)

        def remove_entry():
            if med_type == "regular" and len(self._regular_meds) > 1:
                self._regular_meds.remove(entry_data)
                entry_widget.deleteLater()
                self._update_treatment_text()
            elif med_type == "prn" and len(self._prn_meds) > 1:
                self._prn_meds.remove(entry_data)
                entry_widget.deleteLater()
                self._update_treatment_text()

        med_combo.currentTextChanged.connect(self._update_treatment_text)
        bnf_radio.toggled.connect(self._update_treatment_text)
        above_bnf_radio.toggled.connect(self._update_treatment_text)
        remove_btn.clicked.connect(remove_entry)

    def _update_treatment_text(self):
        """Generate treatment description based on medication entries."""
        regular_meds = []
        prn_meds = []
        above_bnf_meds = []

        # Collect regular medications
        for entry in self._regular_meds:
            name = entry["name"].currentText().strip()
            if name:
                regular_meds.append(name)
                if entry["above_bnf_radio"].isChecked():
                    above_bnf_meds.append(name)

        # Collect PRN medications
        for entry in self._prn_meds:
            name = entry["name"].currentText().strip()
            if name:
                prn_meds.append(name)
                if entry["above_bnf_radio"].isChecked():
                    above_bnf_meds.append(name)

        # Build output text
        parts = []

        if regular_meds:
            parts.append(f"Regular: {', '.join(regular_meds)}")

        if prn_meds:
            parts.append(f"PRN: {', '.join(prn_meds)}")

        if parts:
            text = "; ".join(parts) + "."
            if above_bnf_meds:
                text += f"\n\nAll medication at BNF doses except {', '.join(above_bnf_meds)}."
            else:
                text += "\n\nAll medication at BNF doses."
            self.treatment_desc.setPlainText(text)
        else:
            self.treatment_desc.setPlainText("")

        # Auto-sync to card
        self._update_treatment_card()

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_clinician_card(self):
        parts = []
        if self.approved_clinician_radio.isChecked():
            parts.append("Approved Clinician")
        else:
            parts.append("SOAD")
        if self.clinician_name.text().strip():
            parts.append(self.clinician_name.text().strip())
        if self.clinician_address.text().strip():
            parts.append(self.clinician_address.text().strip())
        self.cards["clinician"].set_content_text("\n".join(parts))

    def _update_patient_card(self):
        parts = []
        if self.patient_name.text().strip():
            parts.append(self.patient_name.text().strip())
        if self.patient_address.text().strip():
            parts.append(self.patient_address.text().strip())
        self.cards["patient"].set_content_text("\n".join(parts) if parts else "Click to enter details")

    def _update_treatment_card(self):
        text = self.treatment_desc.toPlainText()
        if len(text) > 200:
            text = text[:200] + "..."
        self.cards["treatment"].set_content_text(text if text else "Click to enter details")

    def _update_signature_card(self):
        date_str = self.sig_date.date().toString("dd MMM yyyy")
        self.cards["signature"].set_content_text(date_str)

    # ----------------------------------------------------------------
    # HELPERS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("QLineEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 8px; font-size: 18px; } QLineEdit:focus { border-color: #059669; }")
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("QDateEdit { background: white; border: 1px solid #d1d5db; border-radius: 5px; padding: 6px; font-size: 18px; }")
        return date_edit

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.clinician_name.clear()
            self.clinician_address.clear()
            self.clinician_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.treatment_desc.clear()
            self.sig_date.setDate(QDate.currentDate())
            # Clear medication entries
            for entry in self._regular_meds:
                entry["name"].setCurrentText("")
                entry["bnf_radio"].setChecked(True)
            for entry in self._prn_meds:
                entry["name"].setCurrentText("")
                entry["bnf_radio"].setChecked(True)
            # Reset clinician type
            self.approved_clinician_radio.setChecked(True)
            # Clear cards
            for card in self.cards.values():
                card.set_content_text("")
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form T2",
            f"Form_T2_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_T2_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form T2 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color (#918C0D) and cream highlight (#FFFED5)
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)
            CREAM_FILL = 'FFFED5'

            # Pre-process: clean ALL paragraphs - remove permission markers, convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), CREAM_FILL)
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), CREAM_FILL)

            def set_entry_box(para, content: str):
                """Set entry box with gold brackets."""
                if not content or not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Opening bracket
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def format_cream_para(para, content: str):
                """Cream highlight paragraph."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)

            def format_closing_bracket_para(para):
                """Paragraph with cream and closing bracket."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = '                                                                                                    '
                    ct = para.runs[0]
                else:
                    ct = para.add_run('                                                                                                    ')
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                cb = para.add_run(']')
                cb.font.name = 'Arial'
                cb.font.size = Pt(12)
                cb.font.bold = True
                cb.font.color.rgb = BRACKET_COLOR
                rPr_cb = cb._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            paragraphs = doc.paragraphs

            # Clinician details (para 3) - entry box
            clinician_text = self.clinician_name.text().strip()
            if self.clinician_address.text().strip():
                clinician_text += ", " + self.clinician_address.text().strip()
            if self.clinician_email.text().strip():
                clinician_text += ", " + self.clinician_email.text().strip()
            set_entry_box(paragraphs[3], clinician_text)

            # Para 4: Clinician type with bracket structure
            for run in paragraphs[4].runs:
                run.text = ""
            while len(paragraphs[4].runs) > 1:
                paragraphs[4]._element.remove(paragraphs[4].runs[-1]._element)
            pPr4 = paragraphs[4]._element.get_or_add_pPr()
            for old_shd in pPr4.findall(qn('w:shd')):
                pPr4.remove(old_shd)

            # Opening bracket
            if paragraphs[4].runs:
                paragraphs[4].runs[0].text = '['
                ob4 = paragraphs[4].runs[0]
            else:
                ob4 = paragraphs[4].add_run('[')
            ob4.font.name = 'Arial'
            ob4.font.size = Pt(12)
            ob4.font.bold = True
            ob4.font.color.rgb = BRACKET_COLOR
            rPr_ob4 = ob4._element.get_or_add_rPr()
            shd_ob4 = OxmlElement('w:shd')
            shd_ob4.set(qn('w:val'), 'clear')
            shd_ob4.set(qn('w:color'), 'auto')
            shd_ob4.set(qn('w:fill'), CREAM_FILL)
            rPr_ob4.append(shd_ob4)

            # Approved clinician text
            ac_text = "the approved clinician in charge of the treatment described below"
            ac_run = paragraphs[4].add_run(ac_text)
            ac_run.font.name = 'Arial'
            ac_run.font.size = Pt(12)
            rPr_ac = ac_run._element.get_or_add_rPr()
            shd_ac = OxmlElement('w:shd')
            shd_ac.set(qn('w:val'), 'clear')
            shd_ac.set(qn('w:color'), 'auto')
            shd_ac.set(qn('w:fill'), CREAM_FILL)
            rPr_ac.append(shd_ac)
            if not self.approved_clinician_radio.isChecked():
                ac_run.font.strike = True

            # Slash
            slash_run = paragraphs[4].add_run("/")
            slash_run.font.name = 'Arial'
            slash_run.font.size = Pt(12)
            rPr_sl = slash_run._element.get_or_add_rPr()
            shd_sl = OxmlElement('w:shd')
            shd_sl.set(qn('w:val'), 'clear')
            shd_sl.set(qn('w:color'), 'auto')
            shd_sl.set(qn('w:fill'), CREAM_FILL)
            rPr_sl.append(shd_sl)

            # SOAD text
            soad_text = "a registered medical practitioner appointed for the purposes of Part 4 of the Act (a SOAD)"
            soad_run = paragraphs[4].add_run(soad_text)
            soad_run.font.name = 'Arial'
            soad_run.font.size = Pt(12)
            rPr_soad = soad_run._element.get_or_add_rPr()
            shd_soad = OxmlElement('w:shd')
            shd_soad.set(qn('w:val'), 'clear')
            shd_soad.set(qn('w:color'), 'auto')
            shd_soad.set(qn('w:fill'), CREAM_FILL)
            rPr_soad.append(shd_soad)
            if self.approved_clinician_radio.isChecked():
                soad_run.font.strike = True

            # Closing bracket
            cb4 = paragraphs[4].add_run(']')
            cb4.font.name = 'Arial'
            cb4.font.size = Pt(12)
            cb4.font.bold = True
            cb4.font.color.rgb = BRACKET_COLOR
            rPr_cb4 = cb4._element.get_or_add_rPr()
            shd_cb4 = OxmlElement('w:shd')
            shd_cb4.set(qn('w:val'), 'clear')
            shd_cb4.set(qn('w:color'), 'auto')
            shd_cb4.set(qn('w:fill'), CREAM_FILL)
            rPr_cb4.append(shd_cb4)

            # Suffix text
            suf4 = paragraphs[4].add_run('<delete the phrase which does not apply> certify that')
            suf4.font.name = 'Arial'
            suf4.font.size = Pt(12)

            # Patient details (para 6) - entry box
            patient_text = self.patient_name.text().strip()
            if self.patient_address.text().strip():
                patient_text += ", " + self.patient_address.text().strip()
            set_entry_box(paragraphs[6], patient_text)

            # Treatment description (para 8-10) - multi-line box
            treatment_text = self.treatment_desc.toPlainText().strip()
            if not treatment_text:
                treatment_text = '                                                                                                    '

            # Para 8 - opening bracket + content
            for run in paragraphs[8].runs:
                run.text = ""
            while len(paragraphs[8].runs) > 1:
                paragraphs[8]._element.remove(paragraphs[8].runs[-1]._element)
            pPr8 = paragraphs[8]._element.get_or_add_pPr()
            for old_shd in pPr8.findall(qn('w:shd')):
                pPr8.remove(old_shd)
            if paragraphs[8].runs:
                paragraphs[8].runs[0].text = '['
                ob8 = paragraphs[8].runs[0]
            else:
                ob8 = paragraphs[8].add_run('[')
            ob8.font.name = 'Arial'
            ob8.font.size = Pt(12)
            ob8.font.bold = True
            ob8.font.color.rgb = BRACKET_COLOR
            rPr_ob8 = ob8._element.get_or_add_rPr()
            shd_ob8 = OxmlElement('w:shd')
            shd_ob8.set(qn('w:val'), 'clear')
            shd_ob8.set(qn('w:color'), 'auto')
            shd_ob8.set(qn('w:fill'), CREAM_FILL)
            rPr_ob8.append(shd_ob8)
            ct8 = paragraphs[8].add_run(treatment_text)
            ct8.font.name = 'Arial'
            ct8.font.size = Pt(12)
            rPr_ct8 = ct8._element.get_or_add_rPr()
            shd_ct8 = OxmlElement('w:shd')
            shd_ct8.set(qn('w:val'), 'clear')
            shd_ct8.set(qn('w:color'), 'auto')
            shd_ct8.set(qn('w:fill'), CREAM_FILL)
            rPr_ct8.append(shd_ct8)

            # Para 9 - cream continuation
            format_cream_para(paragraphs[9], '                                                                                                    ')

            # Para 10 - cream + closing bracket
            format_closing_bracket_para(paragraphs[10])

            # Para 11: "[If you need to continue...here [ ] and attach...]"
            for run in paragraphs[11].runs:
                run.text = ""
            while len(paragraphs[11].runs) > 1:
                paragraphs[11]._element.remove(paragraphs[11].runs[-1]._element)
            pPr11 = paragraphs[11]._element.get_or_add_pPr()
            for old_shd in pPr11.findall(qn('w:shd')):
                pPr11.remove(old_shd)
            if paragraphs[11].runs:
                paragraphs[11].runs[0].text = '[If you need to continue on a separate sheet please indicate here'
                t11a = paragraphs[11].runs[0]
            else:
                t11a = paragraphs[11].add_run('[If you need to continue on a separate sheet please indicate here')
            t11a.font.name = 'Arial'
            t11a.font.size = Pt(12)
            # Gold bracket open
            ob11 = paragraphs[11].add_run('[')
            ob11.font.bold = True
            ob11.font.color.rgb = BRACKET_COLOR
            rPr_ob11 = ob11._element.get_or_add_rPr()
            shd_ob11 = OxmlElement('w:shd')
            shd_ob11.set(qn('w:val'), 'clear')
            shd_ob11.set(qn('w:color'), 'auto')
            shd_ob11.set(qn('w:fill'), CREAM_FILL)
            rPr_ob11.append(shd_ob11)
            # Placeholder content
            c11 = paragraphs[11].add_run('     ')
            rPr_c11 = c11._element.get_or_add_rPr()
            shd_c11 = OxmlElement('w:shd')
            shd_c11.set(qn('w:val'), 'clear')
            shd_c11.set(qn('w:color'), 'auto')
            shd_c11.set(qn('w:fill'), CREAM_FILL)
            rPr_c11.append(shd_c11)
            # Gold bracket close
            cb11 = paragraphs[11].add_run(']')
            cb11.font.bold = True
            cb11.font.color.rgb = BRACKET_COLOR
            rPr_cb11 = cb11._element.get_or_add_rPr()
            shd_cb11 = OxmlElement('w:shd')
            shd_cb11.set(qn('w:val'), 'clear')
            shd_cb11.set(qn('w:color'), 'auto')
            shd_cb11.set(qn('w:fill'), CREAM_FILL)
            rPr_cb11.append(shd_cb11)
            # Rest of text
            t11b = paragraphs[11].add_run('and attach that sheet to this form.]')
            t11b.font.name = 'Arial'
            t11b.font.size = Pt(12)

            # Para 14: Signed [ ] Date [ ]
            sig_date = self.sig_date.date().toString("dd MMMM yyyy")
            for run in paragraphs[14].runs:
                run.text = ""
            while len(paragraphs[14].runs) > 1:
                paragraphs[14]._element.remove(paragraphs[14].runs[-1]._element)
            pPr14 = paragraphs[14]._element.get_or_add_pPr()
            for old_shd in pPr14.findall(qn('w:shd')):
                pPr14.remove(old_shd)
            if paragraphs[14].runs:
                paragraphs[14].runs[0].text = 'Signed'
                sl14 = paragraphs[14].runs[0]
            else:
                sl14 = paragraphs[14].add_run('Signed')
            sl14.font.name = 'Arial'
            sl14.font.size = Pt(12)
            # First placeholder
            ob14a = paragraphs[14].add_run('[')
            ob14a.font.bold = True
            ob14a.font.color.rgb = BRACKET_COLOR
            rPr_ob14a = ob14a._element.get_or_add_rPr()
            shd_ob14a = OxmlElement('w:shd')
            shd_ob14a.set(qn('w:val'), 'clear')
            shd_ob14a.set(qn('w:color'), 'auto')
            shd_ob14a.set(qn('w:fill'), CREAM_FILL)
            rPr_ob14a.append(shd_ob14a)
            c14a = paragraphs[14].add_run('                              ')
            rPr_c14a = c14a._element.get_or_add_rPr()
            shd_c14a = OxmlElement('w:shd')
            shd_c14a.set(qn('w:val'), 'clear')
            shd_c14a.set(qn('w:color'), 'auto')
            shd_c14a.set(qn('w:fill'), CREAM_FILL)
            rPr_c14a.append(shd_c14a)
            cb14a = paragraphs[14].add_run(']')
            cb14a.font.bold = True
            cb14a.font.color.rgb = BRACKET_COLOR
            rPr_cb14a = cb14a._element.get_or_add_rPr()
            shd_cb14a = OxmlElement('w:shd')
            shd_cb14a.set(qn('w:val'), 'clear')
            shd_cb14a.set(qn('w:color'), 'auto')
            shd_cb14a.set(qn('w:fill'), CREAM_FILL)
            rPr_cb14a.append(shd_cb14a)
            # Date label
            dl14 = paragraphs[14].add_run(' Date')
            dl14.font.name = 'Arial'
            dl14.font.size = Pt(12)
            # Second placeholder
            ob14b = paragraphs[14].add_run('[')
            ob14b.font.bold = True
            ob14b.font.color.rgb = BRACKET_COLOR
            rPr_ob14b = ob14b._element.get_or_add_rPr()
            shd_ob14b = OxmlElement('w:shd')
            shd_ob14b.set(qn('w:val'), 'clear')
            shd_ob14b.set(qn('w:color'), 'auto')
            shd_ob14b.set(qn('w:fill'), CREAM_FILL)
            rPr_ob14b.append(shd_ob14b)
            date_content = sig_date if sig_date else '                              '
            c14b = paragraphs[14].add_run(date_content)
            rPr_c14b = c14b._element.get_or_add_rPr()
            shd_c14b = OxmlElement('w:shd')
            shd_c14b.set(qn('w:val'), 'clear')
            shd_c14b.set(qn('w:color'), 'auto')
            shd_c14b.set(qn('w:fill'), CREAM_FILL)
            rPr_c14b.append(shd_c14b)
            cb14b = paragraphs[14].add_run(']')
            cb14b.font.bold = True
            cb14b.font.color.rgb = BRACKET_COLOR
            rPr_cb14b = cb14b._element.get_or_add_rPr()
            shd_cb14b = OxmlElement('w:shd')
            shd_cb14b.set(qn('w:val'), 'clear')
            shd_cb14b.set(qn('w:color'), 'auto')
            shd_cb14b.set(qn('w:fill'), CREAM_FILL)
            rPr_cb14b.append(shd_cb14b)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form T2 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        # Collect medication data
        regular_meds = []
        for entry in self._regular_meds:
            name = entry["name"].currentText().strip()
            if name:
                regular_meds.append({
                    "name": name,
                    "above_bnf": entry["above_bnf_radio"].isChecked()
                })

        prn_meds = []
        for entry in self._prn_meds:
            name = entry["name"].currentText().strip()
            if name:
                prn_meds.append({
                    "name": name,
                    "above_bnf": entry["above_bnf_radio"].isChecked()
                })

        return {
            "clinician_type": "soad" if self.soad_radio.isChecked() else "approved_clinician",
            "clinician_name": self.clinician_name.text(),
            "clinician_address": self.clinician_address.text(),
            "clinician_email": self.clinician_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "treatment_desc": self.treatment_desc.toPlainText(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
            "regular_meds": regular_meds,
            "prn_meds": prn_meds,
            "cards": {key: card.get_content() for key, card in self.cards.items()},
        }

    def set_state(self, state: dict):
        if not state:
            return
        # Restore clinician type
        if state.get("clinician_type") == "soad":
            self.soad_radio.setChecked(True)
        else:
            self.approved_clinician_radio.setChecked(True)
        self.clinician_name.setText(state.get("clinician_name", ""))
        self.clinician_address.setText(state.get("clinician_address", ""))
        self.clinician_email.setText(state.get("clinician_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.treatment_desc.setPlainText(state.get("treatment_desc", ""))
        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))

        # Restore regular medications
        regular_meds = state.get("regular_meds", [])
        for i, med in enumerate(regular_meds):
            if i >= len(self._regular_meds):
                self._add_medication_entry("regular")
            entry = self._regular_meds[i]
            entry["name"].setCurrentText(med.get("name", ""))
            if med.get("above_bnf"):
                entry["above_bnf_radio"].setChecked(True)
            else:
                entry["bnf_radio"].setChecked(True)

        # Restore PRN medications
        prn_meds = state.get("prn_meds", [])
        for i, med in enumerate(prn_meds):
            if i >= len(self._prn_meds):
                self._add_medication_entry("prn")
            entry = self._prn_meds[i]
            entry["name"].setCurrentText(med.get("name", ""))
            if med.get("above_bnf"):
                entry["above_bnf_radio"].setChecked(True)
            else:
                entry["bnf_radio"].setChecked(True)

        # Restore card contents
        cards_state = state.get("cards", {})
        for key, content in cards_state.items():
            if key in self.cards:
                self.cards[key].set_content_text(content)

        self._update_treatment_text()

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[T2Form] Set patient name: {patient_info['name']}")
