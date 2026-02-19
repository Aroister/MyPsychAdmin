# ============================================================
# DOCX EXPORTER — HTML → DOCX
# ============================================================

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup, NavigableString
import re


class DocxExporter:

    @staticmethod
    def export_html(html: str, output_path: str):
        """
        Accepts an HTML string and writes a DOCX file.
        Handles bold, italic, underline, and proper paragraph structure.
        """
        try:
            # Debug: Show a sample of the HTML to understand the format
            sample = html[:500] if len(html) > 500 else html
            print(f"[DOCX DEBUG] HTML sample: {repr(sample)}")

            # Preprocess: Convert **text** markdown-style bold to <b>text</b>
            # Handle various possible encodings of asterisks
            # 1. Raw asterisks: **text**
            html = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', html)
            # 2. HTML numeric entity: &#42;&#42;text&#42;&#42;
            html = re.sub(r'&#42;&#42;([^&]+?)&#42;&#42;', r'<b>\1</b>', html)
            # 3. HTML named entity (rare): &ast;&ast;text&ast;&ast;
            html = re.sub(r'&ast;&ast;(.+?)&ast;&ast;', r'<b>\1</b>', html)
            # 4. Mixed: handle case where * might appear with entities around
            html = re.sub(r'(?:\*|&#42;|&ast;){2}(.+?)(?:\*|&#42;|&ast;){2}', r'<b>\1</b>', html)

            # Debug: Check if conversion happened
            bold_count = html.count('<b>')
            print(f"[DOCX DEBUG] Bold tags after conversion: {bold_count}")

            doc = Document()
            soup = BeautifulSoup(html, "html.parser")

            def process_element(element, paragraph=None):
                """Recursively process HTML elements into docx paragraphs."""
                if isinstance(element, NavigableString):
                    text = str(element)
                    # Skip empty or whitespace-only strings
                    if not text.strip():
                        return
                    if paragraph is not None:
                        run = paragraph.add_run(text)
                    return

                # Handle block elements - create new paragraph
                if element.name in ['p', 'div']:
                    # Check if this paragraph has meaningful content first
                    text_content = element.get_text(strip=True)
                    if not text_content:
                        return  # Skip empty paragraphs
                    para = doc.add_paragraph()
                    for child in element.children:
                        process_inline(child, para)
                    return

                # Handle br as line break within current context
                if element.name == 'br':
                    return

                # Handle other elements by processing children
                for child in element.children:
                    process_element(child, paragraph)

            def process_inline(element, paragraph):
                """Process inline elements within a paragraph."""
                if isinstance(element, NavigableString):
                    text = str(element)
                    if text.strip() or text == ' ':
                        paragraph.add_run(text)
                    return

                if element.name == 'br':
                    paragraph.add_run('\n')
                    return

                if element.name in ['b', 'strong']:
                    for child in element.children:
                        if isinstance(child, NavigableString):
                            text = str(child)
                            if text.strip() or text == ' ':
                                run = paragraph.add_run(text)
                                run.bold = True
                        else:
                            process_inline(child, paragraph)
                    return

                if element.name in ['i', 'em']:
                    for child in element.children:
                        if isinstance(child, NavigableString):
                            text = str(child)
                            if text.strip() or text == ' ':
                                run = paragraph.add_run(text)
                                run.italic = True
                        else:
                            process_inline(child, paragraph)
                    return

                if element.name == 'u':
                    for child in element.children:
                        if isinstance(child, NavigableString):
                            text = str(child)
                            if text.strip() or text == ' ':
                                run = paragraph.add_run(text)
                                run.underline = True
                        else:
                            process_inline(child, paragraph)
                    return

                # For other inline elements, just process children
                for child in element.children:
                    process_inline(child, paragraph)

            # Process all top-level elements
            for element in soup.children:
                process_element(element)

            # Remove empty paragraphs at the end
            while doc.paragraphs and not doc.paragraphs[-1].text.strip():
                p = doc.paragraphs[-1]._element
                p.getparent().remove(p)

            doc.save(output_path)
            print(f"[DOCX EXPORT] Saved → {output_path}")

        except Exception as e:
            print(f"[DOCX EXPORT ERROR]: {e}")
            import traceback
            traceback.print_exc()
