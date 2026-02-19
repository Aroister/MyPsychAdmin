# ================================================================
#  HCR-20 V3 RISK ASSESSMENT FORM PAGE
#  Historical Clinical Risk Management - 20 Item Scale Version 3
#  Based on HCR-20 V3 (Douglas et al., 2013)
# ================================================================

from __future__ import annotations
import sys
from datetime import datetime
from typing import Optional
from PySide6.QtCore import Qt, Signal, QDate, QEvent
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QFileDialog, QMessageBox, QToolButton,
    QGridLayout, QRadioButton, QButtonGroup, QComboBox, QCompleter,
    QStyleFactory, QSlider, QSizePolicy, QColorDialog,
    QSplitter, QStackedWidget
)
from PySide6.QtGui import QColor, QFontDatabase, QFont, QTextListFormat
from PySide6.QtWidgets import QApplication

# Collapsible and Resizable sections for expandable panels
try:
    from background_history_popup import CollapsibleSection, ResizableSection
except ImportError:
    CollapsibleSection = None
    ResizableSection = None

from utils.resource_path import resource_path
from shared_widgets import create_zoom_row
from mypsy_richtext_editor import MyPsychAdminRichTextEditor

# HCR-20 specific extraction from notes
try:
    from hcr20_extractor import HCR20Extractor
    from shared_data_store import get_shared_store
    HCR20_EXTRACTOR_AVAILABLE = True
except ImportError:
    HCR20_EXTRACTOR_AVAILABLE = False
    HCR20Extractor = None
    get_shared_store = None


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# NO-WHEEL COMBOBOX (prevents scroll from changing value)
# ================================================================
class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# HCR-20 CARD WIDGET
# ================================================================
class HCR20CardWidget(QFrame):
    """A clickable card for an HCR-20 report section."""

    clicked = Signal(str)

    STYLE_NORMAL = """
        HCR20CardWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
        }
        HCR20CardWidget:hover {
            border-color: #1e40af;
            background: #eff6ff;
        }
        QLabel {
            background: transparent;
            border: none;
        }
        QRadioButton, QCheckBox {
            background: transparent;
            border: none;
        }
    """

    STYLE_SELECTED = """
        HCR20CardWidget {
            background: #dbeafe;
            border: 2px solid #1e40af;
            border-left: 4px solid #1e3a8a;
            border-radius: 12px;
        }
        QLabel {
            background: transparent;
            border: none;
        }
        QRadioButton, QCheckBox {
            background: transparent;
            border: none;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
        """)
        header_row.addWidget(title_lbl)
        header_row.addStretch()

        layout.addLayout(header_row)

        # Editor
        self.editor = MyPsychAdminRichTextEditor()
        self.editor.setPlaceholderText("Click to edit...")
        self.editor.setReadOnly(False)
        self._editor_height = 130  # Reduced by 10% (was 144)
        self.editor.setMinimumHeight(72)  # Reduced by 10% (was 80)
        self.editor.setMaximumHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 16px;
                color: #374151;
            }
        """)
        layout.addWidget(self.editor)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.editor, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(9)  # Reduced by 10% (was 10)
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #1e40af;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(60, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setMinimumHeight(self._editor_height)
                self.editor.setMaximumHeight(self._editor_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        return self._selected


class DragResizeBar(QFrame):
    """A thin drag bar that resizes an associated QTextEdit vertically."""

    def __init__(self, text_edit, parent=None):
        super().__init__(parent)
        self._text_edit = text_edit
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        self.setFixedHeight(8)
        self.setCursor(Qt.CursorShape.SizeVerCursor)
        self.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 1px 40px;
            }
            QFrame:hover {
                background: #1e40af;
            }
        """)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._dragging = True
            self._drag_start_y = event.globalPosition().y()
            self._drag_start_height = self._text_edit.height()
            event.accept()

    def mouseMoveEvent(self, event):
        if self._dragging:
            delta = event.globalPosition().y() - self._drag_start_y
            new_height = max(40, min(600, int(self._drag_start_height + delta)))
            self._text_edit.setMinimumHeight(new_height)
            self._text_edit.setMaximumHeight(new_height)
            self._text_edit.setFixedHeight(new_height)
            event.accept()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._dragging = False
            # After drag, set min back to allow future shrink via drag
            current_h = self._text_edit.height()
            self._text_edit.setMinimumHeight(40)
            self._text_edit.setMaximumHeight(16777215)
            self._text_edit.setFixedHeight(current_h)
            event.accept()


# ================================================================
# HCR-20 TOOLBAR (Matching Letter Writer Style)
# ================================================================
class HCR20Toolbar(QWidget):
    """Toolbar for the HCR-20 Form Page with full formatting options."""

    # Formatting signals
    set_font_family = Signal(str)
    set_font_size = Signal(int)
    toggle_bold = Signal()
    toggle_italic = Signal()
    toggle_underline = Signal()
    set_text_color = Signal(QColor)
    set_highlight_color = Signal(QColor)
    set_align_left = Signal()
    set_align_center = Signal()
    set_align_right = Signal()
    set_align_justify = Signal()
    bullet_list = Signal()
    numbered_list = Signal()
    indent = Signal()
    outdent = Signal()
    undo = Signal()
    redo = Signal()
    insert_date = Signal()

    # Action signals
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(58)  # Reduced by 10% (was 64)
        self.setStyleSheet("""
            HCR20Toolbar {
                background: rgba(200, 210, 230, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 500;
            }
            QToolButton:hover {
                background: rgba(0,0,0,0.08);
            }
            QComboBox {
                background: rgba(255,255,255,0.85);
                color: #333333;
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 16px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #e0e0e0;
            }
            QScrollArea {
                background: transparent;
                border: none;
            }
            QScrollBar:horizontal {
                height: 6px;
                background: transparent;
            }
            QScrollBar::handle:horizontal {
                background: rgba(0,0,0,0.2);
                border-radius: 3px;
                min-width: 30px;
            }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0px;
            }
        """)

        # Outer layout for the toolbar
        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        # Scroll area to prevent compression
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(58)  # Reduced by 10% (was 64)

        # Container widget for toolbar content
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(54)  # Reduced by 10% (was 60)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 8, 2)
        layout.setSpacing(10)

        # ---------------------------------------------------------
        # UPLOADED DOCS BUTTON (dropdown menu)
        # ---------------------------------------------------------
        from PySide6.QtWidgets import QMenu
        import_btn = QToolButton()
        import_btn.setText("Uploaded Docs")
        import_btn.setFixedSize(160, 42)
        import_btn.setPopupMode(QToolButton.InstantPopup)
        self.upload_menu = QMenu()
        import_btn.setMenu(self.upload_menu)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 17px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #059669;
            }
            QToolButton:pressed {
                background: #047857;
            }
            QToolButton::menu-indicator { image: none; }
        """)
        import_btn.setToolTip("Select from uploaded documents")
        import_btn.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        layout.addWidget(import_btn)

        # ---------------------------------------------------------
        # FONT FAMILY
        # ---------------------------------------------------------
        self.font_combo = QComboBox()
        self.font_combo.setFixedWidth(180)
        self.font_combo.setFocusPolicy(Qt.FocusPolicy.NoFocus)

        families = QFontDatabase.families()
        if sys.platform == "win32":
            preferred = ["Segoe UI", "Calibri", "Cambria", "Arial", "Times New Roman"]
        else:
            preferred = ["Avenir Next", "Avenir", "SF Pro Text", "Helvetica Neue", "Helvetica"]

        added = set()
        for f in preferred:
            if f in families:
                self.font_combo.addItem(f)
                added.add(f)
        for f in families:
            if f not in added:
                self.font_combo.addItem(f)

        self.font_combo.currentTextChanged.connect(self.set_font_family.emit)
        layout.addWidget(self.font_combo)

        # ---------------------------------------------------------
        # FONT SIZE
        # ---------------------------------------------------------
        self.size_combo = QComboBox()
        self.size_combo.setFixedWidth(65)
        self.size_combo.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        for sz in [8, 9, 10, 11, 12, 14, 16, 18, 20, 22]:
            self.size_combo.addItem(str(sz))
        self.size_combo.setCurrentText("12")
        self.size_combo.currentTextChanged.connect(lambda v: self.set_font_size.emit(int(v)))
        layout.addWidget(self.size_combo)

        # Simple button helper - fixed size to prevent compression
        # NoFocus policy ensures button clicks don't steal focus from text editors
        def btn(label, slot):
            b = QToolButton()
            b.setText(label)
            b.setMinimumWidth(36)
            b.setFocusPolicy(Qt.FocusPolicy.NoFocus)
            b.clicked.connect(slot)
            return b

        # ---------------------------------------------------------
        # BASIC STYLES
        # ---------------------------------------------------------
        layout.addWidget(btn("B", self.toggle_bold.emit))
        layout.addWidget(btn("I", self.toggle_italic.emit))
        layout.addWidget(btn("U", self.toggle_underline.emit))

        # ---------------------------------------------------------
        # COLORS
        # ---------------------------------------------------------
        color_btn = btn("A", self._choose_text_color)
        layout.addWidget(color_btn)

        highlight_btn = btn("🖍", self._choose_highlight_color)
        layout.addWidget(highlight_btn)

        # ---------------------------------------------------------
        # ALIGNMENT
        # ---------------------------------------------------------
        layout.addWidget(btn("L", self.set_align_left.emit))
        layout.addWidget(btn("C", self.set_align_center.emit))
        layout.addWidget(btn("R", self.set_align_right.emit))
        layout.addWidget(btn("J", self.set_align_justify.emit))

        # ---------------------------------------------------------
        # LISTS / INDENTATION
        # ---------------------------------------------------------
        layout.addWidget(btn("•", self.bullet_list.emit))
        layout.addWidget(btn("1.", self.numbered_list.emit))
        layout.addWidget(btn("→", self.indent.emit))
        layout.addWidget(btn("←", self.outdent.emit))

        # ---------------------------------------------------------
        # UNDO / REDO
        # ---------------------------------------------------------
        layout.addWidget(btn("⟲", self.undo.emit))
        layout.addWidget(btn("⟳", self.redo.emit))

        # ---------------------------------------------------------
        # INSERT DATE
        # ---------------------------------------------------------
        layout.addWidget(btn("Date", self.insert_date.emit))

        layout.addStretch()

        # ---------------------------------------------------------
        # EXPORT BUTTON - PROMINENT BLUE STYLING (RIGHT)
        # ---------------------------------------------------------
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(160, 42)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 17px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #1d4ed8;
            }
            QToolButton:pressed {
                background: #1e40af;
            }
        """)
        export_btn.setToolTip("Export to Word document")
        export_btn.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # Finalize scroll area
        scroll.setWidget(container)
        outer_layout.addWidget(scroll)

    # -----------------------------
    # COLOR PICKERS
    # -----------------------------
    def _choose_text_color(self):
        col = QColorDialog.getColor(QColor("black"), self)
        if col.isValid():
            self.set_text_color.emit(col)

    def _choose_highlight_color(self):
        col = QColorDialog.getColor(QColor("yellow"), self)
        if col.isValid():
            self.set_highlight_color.emit(col)


# ================================================================
# HCR-20 V3 FORM PAGE
# ================================================================
class HCR20FormPage(QWidget):
    """HCR-20 V3 Risk Assessment Form Page with card/popup interface."""

    go_back = Signal()

    # HCR-20 V3 Historical Items (H1-H10)
    HISTORICAL_ITEMS = [
        ("H1", "History of Problems with Violence"),
        ("H2", "History of Problems with Other Antisocial Behaviour"),
        ("H3", "History of Problems with Relationships"),
        ("H4", "History of Problems with Employment"),
        ("H5", "History of Problems with Substance Use"),
        ("H6", "History of Problems with Major Mental Disorder"),
        ("H7", "History of Problems with Personality Disorder"),
        ("H8", "History of Problems with Traumatic Experiences"),
        ("H9", "History of Problems with Violent Attitudes"),
        ("H10", "History of Problems with Treatment or Supervision Response"),
    ]

    # HCR-20 V3 Clinical Items (C1-C5)
    CLINICAL_ITEMS = [
        ("C1", "Recent Problems with Insight"),
        ("C2", "Recent Problems with Violent Ideation or Intent"),
        ("C3", "Recent Problems with Symptoms of Major Mental Disorder"),
        ("C4", "Recent Problems with Instability"),
        ("C5", "Recent Problems with Treatment or Supervision Response"),
    ]

    # HCR-20 V3 Risk Management Items (R1-R5)
    RISK_ITEMS = [
        ("R1", "Future Problems with Professional Services and Plans"),
        ("R2", "Future Problems with Living Situation"),
        ("R3", "Future Problems with Personal Support"),
        ("R4", "Future Problems with Treatment or Supervision Response"),
        ("R5", "Future Problems with Stress or Coping"),
    ]

    # All sections for the form (HCR-20 V3)
    SECTIONS = [
        ("Patient Details", "patient_details"),
        ("Assessment Details", "assessment_details"),
        ("Sources of Information", "sources"),
        # Historical Items (H1-H10)
        ("H1. History of Problems with Violence", "h1"),
        ("H2. History of Problems with Other Antisocial Behaviour", "h2"),
        ("H3. History of Problems with Relationships", "h3"),
        ("H4. History of Problems with Employment", "h4"),
        ("H5. History of Problems with Substance Use", "h5"),
        ("H6. History of Problems with Major Mental Disorder", "h6"),
        ("H7. History of Problems with Personality Disorder", "h7"),
        ("H8. History of Problems with Traumatic Experiences", "h8"),
        ("H9. History of Problems with Violent Attitudes", "h9"),
        ("H10. History of Problems with Treatment or Supervision Response", "h10"),
        # Clinical Items (C1-C5) - Rated over the last six months
        ("C1. Recent Problems with Insight", "c1"),
        ("C2. Recent Problems with Violent Ideation or Intent", "c2"),
        ("C3. Recent Problems with Symptoms of Major Mental Disorder", "c3"),
        ("C4. Recent Problems with Instability", "c4"),
        ("C5. Recent Problems with Treatment or Supervision Response", "c5"),
        # Risk Management Items (R1-R5) - Considered over the next six months
        ("R1. Future Problems with Professional Services and Plans", "r1"),
        ("R2. Future Problems with Living Situation", "r2"),
        ("R3. Future Problems with Personal Support", "r3"),
        ("R4. Future Problems with Treatment or Supervision Response", "r4"),
        ("R5. Future Problems with Stress or Coping", "r5"),
        # Formulation and Scenarios
        ("Violence Risk Formulation", "formulation"),
        ("Scenarios - Nature", "scenario_nature"),
        ("Scenarios - Severity", "scenario_severity"),
        ("Scenarios - Imminence", "scenario_imminence"),
        ("Scenarios - Frequency", "scenario_frequency"),
        ("Scenarios - Likelihood", "scenario_likelihood"),
        # Management
        ("Risk Enhancing Factors", "risk_enhancing"),
        ("Protective Factors", "protective"),
        ("Monitoring", "monitoring"),
        ("Treatment Recommendations", "treatment"),
        ("Supervision", "supervision"),
        ("Victim Safety Planning", "victim_safety"),
        # Signature
        ("Signature", "signature"),
    ]

    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self._last_generated_text = {}
        self._current_gender = "male"
        self._my_details = self._load_my_details()
        self._setup_ui()

    def _load_my_details(self) -> dict:
        """Load clinician details from database."""
        if not self.db:
            return {}

        details = self.db.get_clinician_details()
        if not details:
            return {}

        # details is a tuple: (id, full_name, role_title, discipline, registration_body,
        #                      registration_number, phone, email, team_service,
        #                      hospital_org, ward_department, signature_block)
        return {
            "full_name": details[1] or "",
            "role_title": details[2] or "",
            "discipline": details[3] or "",
            "registration_body": details[4] or "",
            "registration_number": details[5] or "",
            "phone": details[6] or "",
            "email": details[7] or "",
            "team_service": details[8] or "",
            "hospital_org": details[9] or "",
            "ward_department": details[10] or "",
            "signature_block": details[11] or "",
        }

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header with Clear Form button on right
        header = QFrame()
        header.setFixedHeight(43)  # Reduced by 10% (was 48)
        header.setStyleSheet("background: #1e40af; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("HCR-20 V3 Risk Assessment Report")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Clear Form button in header (right side)
        clear_btn = QPushButton("Clear Form")
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setFixedSize(220, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #1e40af;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover {
                background: #1e3a8a;
            }
            QPushButton:pressed {
                background: #172554;
            }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Full Toolbar with formatting options
        self.toolbar = HCR20Toolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())

        # Connect formatting signals
        self.toolbar.set_font_family.connect(self._set_font_family)
        self.toolbar.set_font_size.connect(self._set_font_size)
        self.toolbar.toggle_bold.connect(self._toggle_bold)
        self.toolbar.toggle_italic.connect(self._toggle_italic)
        self.toolbar.toggle_underline.connect(self._toggle_underline)
        self.toolbar.set_text_color.connect(self._set_text_color)
        self.toolbar.set_highlight_color.connect(self._set_highlight_color)
        self.toolbar.set_align_left.connect(self._set_align_left)
        self.toolbar.set_align_center.connect(self._set_align_center)
        self.toolbar.set_align_right.connect(self._set_align_right)
        self.toolbar.set_align_justify.connect(self._set_align_justify)
        self.toolbar.bullet_list.connect(self._toggle_bullet_list)
        self.toolbar.numbered_list.connect(self._toggle_numbered_list)
        self.toolbar.indent.connect(self._indent)
        self.toolbar.outdent.connect(self._outdent)
        self.toolbar.undo.connect(self._undo)
        self.toolbar.redo.connect(self._redo)
        self.toolbar.insert_date.connect(self._insert_date)

        main_layout.addWidget(self.toolbar)

        # Track the currently focused text editor for toolbar actions
        self._active_editor = None

        # HCR-20 Extractor instance
        self._hcr20_extractor = HCR20Extractor() if HCR20_EXTRACTOR_AVAILABLE else None

        # Store raw notes for imported data
        self._extracted_raw_notes = []

        # Connect to SharedDataStore for notes
        self._connect_shared_store()

        # Content area with splitter
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards scroll area
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setStyleSheet("""
            QScrollArea {
                background: #f3f4f6;
                border: none;
            }
        """)
        self.main_splitter.addWidget(self.cards_holder)

        self.cards_root = QWidget()
        self.cards_root.setStyleSheet("background: #f3f4f6;")
        self.cards_layout = QVBoxLayout(self.cards_root)
        self.cards_layout.setContentsMargins(32, 24, 32, 24)
        self.cards_layout.setSpacing(16)
        self.cards_holder.setWidget(self.cards_root)

        # Right: Panel with popup stack
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(400)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame {
                background: rgba(245,245,245,0.98);
                border-left: 1px solid rgba(0,0,0,0.08);
            }
        """)
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 500])

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 22px;
            font-weight: 700;
            color: #1e40af;
            background: rgba(30, 64, 175, 0.1);
            padding: 8px 12px;
            border-radius: 8px;
        """)
        panel_layout.addWidget(self.panel_title)

        # Popup stack
        self.popup_stack = QStackedWidget()
        self.popup_stack.setStyleSheet("background: white; border-radius: 8px;")
        self.popup_stack.setMinimumHeight(144)  # Reduced by 10% (was 160)

        # Add a placeholder widget
        placeholder = QLabel("Click a section card to edit")
        placeholder.setAlignment(Qt.AlignmentFlag.AlignCenter)
        placeholder.setStyleSheet("font-size: 18px; color: #6b7280; background: white;")
        self.popup_stack.addWidget(placeholder)

        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        # Initialize cards and popups
        self.cards = {}
        self.popups = {}
        self.popup_send_buttons = {}
        self.popup_generators = {}
        self._selected_card_key = None
        self._syncing = False  # Flag to prevent infinite sync loops

        # Create all cards
        self._create_cards()

        # Build popup content
        self._build_popups()

        # Connect card editors to sync back to popups
        self._connect_card_to_popup_sync()

        # Set up editor focus tracking for toolbar formatting
        self._setup_editor_tracking()

    def _create_cards(self):
        """Create all section cards with section headers."""
        # Add section headers for organization
        sections_by_category = {
            "header": ["patient_details", "assessment_details", "sources"],
            "historical": ["h1", "h2", "h3", "h4", "h5", "h6", "h7", "h8", "h9", "h10"],
            "clinical": ["c1", "c2", "c3", "c4", "c5"],
            "risk": ["r1", "r2", "r3", "r4", "r5"],
            "formulation": ["formulation", "scenario_nature", "scenario_severity", "scenario_imminence", "scenario_frequency", "scenario_likelihood"],
            "management": ["risk_enhancing", "protective", "monitoring", "treatment", "supervision", "victim_safety"],
            "signature": ["signature"],
        }

        category_labels = {
            "header": None,
            "historical": "Historical Items (H1-H10)",
            "clinical": "Clinical Items (C1-C5)",
            "risk": "Risk Management Items (R1-R5)",
            "formulation": "Violence Risk Formulation & Scenarios",
            "management": "Proposed Management Strategies",
            "signature": None,
        }

        current_category = None

        for title, key in self.SECTIONS:
            # Find which category this key belongs to
            for cat, keys in sections_by_category.items():
                if key in keys:
                    if cat != current_category:
                        current_category = cat
                        label_text = category_labels.get(cat)
                        if label_text:
                            header_lbl = QLabel(label_text)
                            header_lbl.setStyleSheet("""
                                font-size: 20px;
                                font-weight: 700;
                                color: #1e40af;
                                padding: 16px 0 8px 0;
                                background: transparent;
                            """)
                            self.cards_layout.addWidget(header_lbl)
                    break

            card = HCR20CardWidget(title, key, parent=self.cards_root)
            card.clicked.connect(self._on_card_clicked)
            self.cards[key] = card
            self.cards_layout.addWidget(card)

        self.cards_layout.addStretch()

    def _on_card_clicked(self, key: str):
        """Handle card click - show appropriate popup."""
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        # Deselect previous
        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)

        # Select new
        self._selected_card_key = key
        if key in self.cards:
            self.cards[key].setSelected(True)

        # Show popup
        if key in self.popups:
            self.popup_stack.setCurrentWidget(self.popups[key])

        # Auto-populate risk formulation sections from H1-R5 items when navigating to them
        if key == "risk_enhancing":
            self._auto_populate_risk_formulation()

    def _build_popups(self):
        """Build all section popups with their controls."""
        self._build_popup_patient_details()
        self._build_popup_assessment_details()
        self._build_popup_sources()

        # Define subsections for each HCR-20 V3 item based on the JT HCR-20 document structure
        HCR_SUBSECTIONS = {
            # Historical Items
            "H1": [
                ("Child (aged 12 and under):", "Document any violence or aggressive behaviour during childhood. Consider fights, bullying, cruelty to animals, destruction of property, fire setting, or other violent acts before age 13."),
                ("Adolescent (aged 13-17):", "Document violence during adolescence. Consider physical fights, assault, weapon use, threats, bullying, property destruction, or involvement with violent groups during teenage years."),
                ("Adult (aged 18+):", "Document adult violence history. Include physical assaults, domestic violence, violence toward staff/professionals, weapons use, threats, and any criminal convictions for violent offences."),
            ],
            "H2": [
                ("Child (aged 12 and under):", "Document antisocial behaviour during childhood. Consider theft, vandalism, truancy, lying, rule-breaking, running away, and other conduct problems before age 13."),
                ("Adolescent (aged 13-17):", "Document antisocial behaviour during adolescence. Include theft, fraud, drug dealing, vandalism, arson (without violence), sexual offences, and other criminal or antisocial conduct."),
                ("Adult (aged 18+):", "Document adult antisocial behaviour. Consider criminal convictions, probation/parole violations, fraud, property offences, and persistent rule-breaking or irresponsible behaviour."),
            ],
            "H3": [
                ("Intimate Relationships:", "Document history of intimate relationships. Consider stability, conflict, domestic abuse (as perpetrator or victim), separation patterns, and quality of romantic partnerships."),
                ("Non-intimate Relationships:", "Document relationships with family, friends, and colleagues. Consider social isolation, interpersonal conflict, exploitation of others, and ability to maintain supportive relationships."),
            ],
            "H4": [
                ("Education:", "Document educational history. Include academic achievement, behavioural problems at school, truancy, suspensions/expulsions, special educational needs, and highest qualification achieved."),
                ("Employment:", "Document employment history. Consider job stability, reasons for job losses, workplace conflicts, longest period of employment, and current vocational status."),
            ],
            "H5": [
                ("Substance Use History:", "Document history of alcohol and drug use. Include age of first use, substances used, patterns of use, periods of abstinence, and relationship between substance use and violent behaviour."),
                ("Treatment History:", "Document any substance misuse treatment. Include detoxification, rehabilitation programmes, AA/NA attendance, and outcomes of previous treatment attempts."),
                ("Current Status:", "Document current substance use status including recent use, current abstinence, relapse patterns, and engagement with recovery support."),
            ],
            "H6": [
                ("General:", "Document history of major mental disorder. Include age of onset, course of illness, periods of remission, and overall impact on functioning."),
                ("Psychotic Disorders:", "Document history of psychotic symptoms including schizophrenia, schizoaffective disorder, delusional disorder. Include nature of delusions/hallucinations and relationship to violence."),
                ("Major Mood Disorders:", "Document history of major depressive disorder, bipolar disorder, or other major mood disorders. Include manic episodes, depressive episodes, and relationship to risk behaviours."),
                ("Other Mental Disorders:", "Document other significant mental health conditions including anxiety disorders, PTSD, OCD, eating disorders, and their impact on risk."),
            ],
            "H7": [
                ("Personality Disorder Features:", "Document history of personality disorder or features. Include formal diagnoses, personality assessments (e.g., PCL-R), and relevant traits such as antisocial, borderline, narcissistic, or paranoid features."),
                ("Impact on Functioning:", "Document how personality difficulties have impacted relationships, employment, treatment engagement, and violent behaviour."),
            ],
            "H8": [
                ("Victimization/Trauma:", "Document history of victimization including physical abuse, sexual abuse, emotional abuse, neglect, bullying, or witnessing domestic violence. Include age at time of trauma and impact."),
                ("Adverse Childrearing Experiences:", "Document adverse childhood experiences including parental separation, parental mental illness, parental substance misuse, parental criminality, poverty, and unstable care arrangements."),
            ],
            "H9": [
                ("Violent Attitudes:", "Document attitudes supportive of violence. Consider beliefs that violence is acceptable or justified, lack of remorse for past violence, positive views about aggression, and attitudes toward specific victim groups."),
                ("Antisocial Attitudes:", "Document broader antisocial attitudes including criminal thinking patterns, disregard for rules/laws, lack of empathy, and sense of entitlement."),
            ],
            "H10": [
                ("Treatment Response:", "Document history of response to treatment. Include engagement with mental health services, medication compliance, participation in psychological interventions, and outcomes."),
                ("Supervision Response:", "Document response to supervision. Include compliance with probation/parole, hospital leave conditions, ward rules, and any breaches or failures."),
            ],
            # Clinical Items (rated over the last six months)
            "C1": [
                ("Insight into Mental Health:", "Does the individual recognise and accept their mental health diagnosis? Do they understand the nature and impact of their mental illness?"),
                ("Insight into Violence Risk:", "Does the individual recognise their risk factors for violence? Do they understand what triggers their violent behaviour and accept responsibility for past violence?"),
                ("Insight into Need for Treatment:", "Does the individual accept the need for treatment? Do they understand why treatment is necessary and show willingness to engage?"),
            ],
            "C2": [
                ("Violent Ideation:", "Document current thoughts about violence including fantasies, plans, intentions, or preoccupations with violence. Consider frequency, intensity, and specificity of violent thoughts."),
                ("Violent Intent:", "Document any current stated or implied intent to harm others. Consider threats made, identified targets, and any planning or preparatory behaviours."),
            ],
            "C3": [
                ("Symptoms of Psychotic Disorders:", "Document current psychotic symptoms including delusions, hallucinations (especially command hallucinations), paranoid ideation, and disorganised thinking. Rate severity and relationship to violence risk."),
                ("Symptoms of Major Mood Disorder:", "Document current mood symptoms including depression, mania, mixed states, irritability, and emotional dysregulation. Rate severity and relationship to violence risk."),
            ],
            "C4": [
                ("Affective Instability:", "Document emotional instability including mood swings, irritability, anger outbursts, and difficulty regulating emotions over the past six months."),
                ("Behavioural Instability:", "Document behavioural instability including impulsive actions, risk-taking, self-harm, aggression, and difficulty maintaining consistent behaviour patterns."),
                ("Cognitive Instability:", "Document cognitive instability including concentration difficulties, confusion, disorientation, or rapidly changing beliefs and perceptions."),
            ],
            "C5": [
                ("Treatment Engagement:", "Document current engagement with treatment including attendance at appointments, participation in therapy, and relationship with treatment providers."),
                ("Medication Compliance:", "Document current medication compliance including taking medication as prescribed, attitudes toward medication, and any recent non-compliance."),
                ("Supervision Compliance:", "Document current compliance with supervision requirements including ward rules, leave conditions, and any recent breaches."),
            ],
            # Risk Management Items (considered over the next six months)
            "R1": [
                ("Hospital:", "If remaining in hospital, what professional services and plans are in place? Consider care planning, treatment programmes, multidisciplinary input, and adequacy of current arrangements."),
                ("Community:", "If discharged to community, what professional services and plans would be needed? Consider CPA arrangements, community mental health support, and feasibility of proposed plans."),
            ],
            "R2": [
                ("Hospital:", "If remaining in hospital, what is the living situation? Consider ward environment, peer group, level of security, and any environmental risk factors."),
                ("Community:", "If discharged to community, what would the living situation be? Consider housing stability, neighbourhood, proximity to victims or antisocial peers, and access to weapons/substances."),
            ],
            "R3": [
                ("Hospital:", "If remaining in hospital, what personal support is available? Consider family contact, peer relationships on ward, and therapeutic relationships with staff."),
                ("Community:", "If discharged to community, what personal support would be available? Consider family relationships, friendships, support networks, and potential for isolation."),
            ],
            "R4": [
                ("Hospital:", "If remaining in hospital, how likely is compliance with treatment and supervision? Consider current engagement, motivation, and barriers to compliance."),
                ("Community:", "If discharged to community, how likely is compliance with treatment and supervision? Consider history of community compliance and motivation for ongoing engagement."),
            ],
            "R5": [
                ("Hospital:", "If remaining in hospital, what stressors might be encountered? Consider conflicts with peers/staff, boredom, family issues, and coping resources available."),
                ("Community:", "If discharged to community, what stressors might be encountered? Consider housing, finances, relationships, employment, and coping capacity."),
            ],
        }

        # Build Historical Items
        for code, desc in self.HISTORICAL_ITEMS:
            if code == "H3":
                # H3 has special handling with separate imported data sections for each relationship type
                self._build_hcr_h3_popup(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            elif code == "H4":
                # H4 has special handling with separate imported data sections for Education and Employment
                self._build_hcr_h4_popup(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            elif code == "H7":
                # H7 has special handling with personality disorder features
                self._build_hcr_h7_popup(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            elif code == "H8":
                # H8 has special handling with trauma categories input section
                self._build_hcr_h8_popup(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            elif code == "H9":
                # H9 has special handling with violent attitudes input section
                self._build_hcr_h9_popup(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            elif code == "H10":
                # H10 has special handling with treatment/supervision response categories
                self._build_hcr_h10_popup(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            elif code in HCR_SUBSECTIONS:
                self._build_hcr_item_popup_with_subsections(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            else:
                self._build_hcr_item_popup(code.lower(), code, desc, "historical")

        # Build Clinical Items
        for code, desc in self.CLINICAL_ITEMS:
            if code == "C1":
                # C1 has special handling with insight categories
                self._build_hcr_c1_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "C2":
                # C2 has special handling with violent ideation categories
                self._build_hcr_c2_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "C3":
                # C3 has special handling with symptoms categories
                self._build_hcr_c3_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "C4":
                # C4 has special handling with instability categories
                self._build_hcr_c4_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "C5":
                # C5 has special handling with treatment response categories
                self._build_hcr_c5_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code in HCR_SUBSECTIONS:
                self._build_hcr_item_popup_with_subsections(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            else:
                self._build_hcr_item_popup(code.lower(), code, desc, "clinical")

        # Build Risk Management Items
        for code, desc in self.RISK_ITEMS:
            if code == "R1":
                self._build_hcr_r1_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "R2":
                self._build_hcr_r2_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "R3":
                self._build_hcr_r3_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "R4":
                self._build_hcr_r4_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code == "R5":
                self._build_hcr_r5_popup(code.lower(), code, desc, HCR_SUBSECTIONS.get(code, []))
            elif code in HCR_SUBSECTIONS:
                self._build_hcr_item_popup_with_subsections(code.lower(), code, desc, HCR_SUBSECTIONS[code])
            else:
                self._build_hcr_item_popup(code.lower(), code, desc, "risk")

        # Formulation and Scenarios - Enhanced with tick boxes and auto-population
        self._build_popup_formulation()

        # Scenario sections (1-5)
        self._build_popup_nature_of_risk()      # 1. Nature of Risk
        self._build_popup_severity()            # 2. Severity
        self._build_popup_imminence()           # 3. Imminence
        self._build_popup_frequency()           # 4. Frequency
        self._build_popup_likelihood()          # 5. Likelihood

        # Management sections (6-11)
        self._build_popup_risk_enhancing()      # 6. Risk-Enhancing Factors
        self._build_popup_protective_factors()  # 7. Protective Factors
        self._build_popup_monitoring()          # 8. Risk Monitoring Indicators
        self._build_popup_management()          # 9. Risk Management Strategies
        self._build_popup_supervision_recs()    # 10. Supervision Recommendations
        self._build_popup_victim_safety()       # 11. Victim Safety Planning

        self._build_popup_signature()

    def _create_popup_container(self, key: str) -> tuple:
        """Create popup with input fields only (auto-syncs to card on change)."""
        main_widget = QWidget()
        main_widget.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(main_widget)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(0)

        input_scroll = QScrollArea()
        input_scroll.setWidgetResizable(True)
        input_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        input_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        input_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        input_container = QWidget()
        input_container.setStyleSheet("background: transparent;")
        input_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        input_layout = QVBoxLayout(input_container)
        input_layout.setContentsMargins(12, 12, 12, 12)
        input_layout.setSpacing(10)

        input_scroll.setWidget(input_container)
        main_layout.addWidget(input_scroll)

        self.popups[key] = main_widget
        self.popup_stack.addWidget(main_widget)

        return input_container, input_layout

    def _add_send_button(self, layout, key: str, generate_func):
        """Register generator function and trigger initial card sync."""
        self.popup_generators[key] = generate_func
        self._update_preview(key)

    def _update_preview(self, key: str):
        """Auto-sync generated text to card."""
        if self._syncing:
            return  # Prevent infinite loop

        if key in self.popup_generators:
            try:
                self._syncing = True
                new_generated_text = self.popup_generators[key]()
                if key in self.cards:
                    current_text = self.cards[key].editor.toPlainText()
                    last_generated = self._last_generated_text.get(key, "")

                    if not current_text or current_text == last_generated:
                        self.cards[key].editor.setPlainText(new_generated_text if new_generated_text else "")
                    elif last_generated and last_generated in current_text:
                        idx = current_text.find(last_generated)
                        if idx == 0:
                            user_additions = current_text[len(last_generated):]
                            self.cards[key].editor.setPlainText((new_generated_text or "") + user_additions)
                        else:
                            before = current_text[:idx]
                            after = current_text[idx + len(last_generated):]
                            self.cards[key].editor.setPlainText(before + (new_generated_text or "") + after)
                    else:
                        if new_generated_text and new_generated_text not in current_text:
                            self.cards[key].editor.setPlainText(new_generated_text if new_generated_text else "")

                    self._last_generated_text[key] = new_generated_text or ""
            except Exception:
                pass  # Silent fail for preview updates
            finally:
                self._syncing = False

    def _sync_card_to_popup(self, key: str):
        """Sync card editor content back to popup fields."""
        if self._syncing:
            return  # Prevent infinite loop

        if key not in self.cards:
            return

        # Skip cards with complex multi-field popups that can't be reverse-synced
        skip_keys = ['patient_details', 'assessment_details', 'signature']
        if key in skip_keys:
            return

        try:
            self._syncing = True
            card_text = self.cards[key].editor.toPlainText()

            # Try multiple possible field naming patterns
            field_patterns = [
                f"popup_{key}_evidence",      # e.g., popup_c3_evidence
                f"popup_{key}_text",          # e.g., popup_sources_text, popup_formulation_text
                f"popup_{key}",               # e.g., popup_formulation (unlikely but check)
            ]

            for attr_name in field_patterns:
                if hasattr(self, attr_name):
                    widget = getattr(self, attr_name)
                    if hasattr(widget, 'toPlainText') and hasattr(widget, 'setPlainText'):
                        if widget.toPlainText() != card_text:
                            widget.setPlainText(card_text)
                        return

            # Special case for formulation
            if key == 'formulation' and hasattr(self, 'popup_formulation_text'):
                if self.popup_formulation_text.toPlainText() != card_text:
                    self.popup_formulation_text.setPlainText(card_text)
                return

            # For items WITH subsections, sync to the first subsection field
            # (We can't easily split back into subsections from combined text)
            for attr_name in dir(self):
                if attr_name.startswith(f"popup_{key}_") and not any(x in attr_name for x in ['presence', 'relevance', 'imported', 'checkboxes', 'group', 'section', 'entries', 'layout']):
                    widget = getattr(self, attr_name, None)
                    if widget and hasattr(widget, 'toPlainText') and hasattr(widget, 'setPlainText'):
                        # Only update if content is different
                        if widget.toPlainText() != card_text:
                            widget.setPlainText(card_text)
                        return  # Only sync to first subsection field

        finally:
            self._syncing = False

    def _connect_preview_updates(self, key: str, widgets: list):
        """Connect widgets to trigger preview updates when changed."""
        for widget in widgets:
            if isinstance(widget, QLineEdit):
                widget.textChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QTextEdit):
                widget.textChanged.connect(lambda k=key: self._update_preview(k))
            elif isinstance(widget, QComboBox):
                widget.currentTextChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QCheckBox):
                widget.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QRadioButton):
                widget.toggled.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QDateEdit):
                widget.dateChanged.connect(lambda _, k=key: self._update_preview(k))
            elif isinstance(widget, QSlider):
                widget.valueChanged.connect(lambda _, k=key: self._update_preview(k))

    def _connect_card_to_popup_sync(self):
        """Connect card editors to sync their content back to popup fields."""
        for key, card in self.cards.items():
            if hasattr(card, 'editor'):
                # Use a lambda with default argument to capture key correctly
                card.editor.textChanged.connect(lambda k=key: self._sync_card_to_popup(k))

    # ================================================================
    # POPUP BUILDERS
    # ================================================================

    def _build_popup_patient_details(self):
        """Build patient details popup."""
        container, layout = self._create_popup_container("patient_details")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 17px; }"

        # Patient name
        name_lbl = QLabel("Patient Name:")
        name_lbl.setStyleSheet(label_style)
        layout.addWidget(name_lbl)
        self.popup_patient_name = QLineEdit()
        self.popup_patient_name.setStyleSheet(input_style)
        layout.addWidget(self.popup_patient_name)

        # Gender
        gender_lbl = QLabel("Gender:")
        gender_lbl.setStyleSheet(label_style)
        layout.addWidget(gender_lbl)
        gender_row = QHBoxLayout()
        self.popup_gender_group = QButtonGroup(self)
        self.popup_gender_male = QRadioButton("Male")
        self.popup_gender_male.setStyleSheet(radio_style)
        self.popup_gender_female = QRadioButton("Female")
        self.popup_gender_female.setStyleSheet(radio_style)
        self.popup_gender_group.addButton(self.popup_gender_male)
        self.popup_gender_group.addButton(self.popup_gender_female)
        self.popup_gender_male.setChecked(True)
        gender_row.addWidget(self.popup_gender_male)
        gender_row.addWidget(self.popup_gender_female)
        gender_row.addStretch()
        layout.addLayout(gender_row)

        # DOB
        dob_lbl = QLabel("Date of Birth:")
        dob_lbl.setStyleSheet(label_style)
        layout.addWidget(dob_lbl)
        self.popup_dob = QDateEdit()
        self.popup_dob.setCalendarPopup(True)
        self.popup_dob.setDisplayFormat("dd/MM/yyyy")
        self.popup_dob.setDate(QDate(1990, 1, 1))
        self.popup_dob.setStyleSheet(input_style)
        layout.addWidget(self.popup_dob)

        # NHS Number
        nhs_lbl = QLabel("NHS Number:")
        nhs_lbl.setStyleSheet(label_style)
        layout.addWidget(nhs_lbl)
        self.popup_nhs = QLineEdit()
        self.popup_nhs.setStyleSheet(input_style)
        layout.addWidget(self.popup_nhs)

        # Address
        addr_lbl = QLabel("Address:")
        addr_lbl.setStyleSheet(label_style)
        layout.addWidget(addr_lbl)
        self.popup_address = QTextEdit()
        self.popup_address.setMaximumHeight(58)  # Reduced by 10% (was 64)
        self.popup_address.setStyleSheet(input_style)
        layout.addWidget(self.popup_address)

        # Date of Admission
        admission_lbl = QLabel("Date of Admission:")
        admission_lbl.setStyleSheet(label_style)
        layout.addWidget(admission_lbl)
        self.popup_admission_date = QDateEdit()
        self.popup_admission_date.setCalendarPopup(True)
        self.popup_admission_date.setDisplayFormat("dd/MM/yyyy")
        self.popup_admission_date.setDate(QDate.currentDate())
        self.popup_admission_date.setStyleSheet(input_style)
        layout.addWidget(self.popup_admission_date)

        # Legal Status
        legal_lbl = QLabel("Legal Status:")
        legal_lbl.setStyleSheet(label_style)
        layout.addWidget(legal_lbl)
        self.popup_legal_status = QLineEdit()
        self.popup_legal_status.setPlaceholderText("e.g., Section 3 of the Mental Health Act 1983")
        self.popup_legal_status.setStyleSheet(input_style)
        layout.addWidget(self.popup_legal_status)

        layout.addStretch()

        def generate():
            parts = []
            name = self.popup_patient_name.text().strip()
            if name:
                parts.append(f"NAME: {name}")
            dob = self.popup_dob.date().toString("d MMMM yyyy")
            parts.append(f"D.O.B: {dob}")
            # Calculate age
            today = QDate.currentDate()
            age = today.year() - self.popup_dob.date().year()
            if today.month() < self.popup_dob.date().month() or (today.month() == self.popup_dob.date().month() and today.day() < self.popup_dob.date().day()):
                age -= 1
            parts.append(f"AGE: {age}")
            nhs = self.popup_nhs.text().strip()
            if nhs:
                parts.append(f"NHS NUMBER: {nhs}")
            addr = self.popup_address.toPlainText().strip()
            if addr:
                parts.append(f"ADDRESS: {addr}")
            admission = self.popup_admission_date.date().toString("d MMMM yyyy")
            parts.append(f"DATE OF ADMISSION: {admission}")
            legal = self.popup_legal_status.text().strip()
            if legal:
                parts.append(f"LEGAL STATUS: {legal}")
            return "\n".join(parts)

        self._connect_preview_updates("patient_details", [
            self.popup_patient_name, self.popup_gender_male, self.popup_gender_female,
            self.popup_dob, self.popup_nhs, self.popup_address,
            self.popup_admission_date, self.popup_legal_status
        ])
        self._add_send_button(layout, "patient_details", generate)

        # Update gender when changed
        self.popup_gender_male.toggled.connect(self._on_gender_changed)
        self.popup_gender_female.toggled.connect(self._on_gender_changed)

    def _on_gender_changed(self):
        """Update current gender for pronoun generation."""
        if self.popup_gender_male.isChecked():
            self._current_gender = "male"
        else:
            self._current_gender = "female"

    def _build_popup_assessment_details(self):
        """Build assessment details popup."""
        container, layout = self._create_popup_container("assessment_details")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        # Tools included
        tools_lbl = QLabel("Structured Clinical Judgment Tools:")
        tools_lbl.setStyleSheet(label_style)
        layout.addWidget(tools_lbl)
        self.popup_tools = QLineEdit()
        self.popup_tools.setText("HCR-20 V3")
        self.popup_tools.setStyleSheet(input_style)
        layout.addWidget(self.popup_tools)

        # Author of original report - auto-fill from MyDetails
        author_orig_lbl = QLabel("Author of Original Report:")
        author_orig_lbl.setStyleSheet(label_style)
        layout.addWidget(author_orig_lbl)
        self.popup_author_original = QLineEdit()
        self.popup_author_original.setStyleSheet(input_style)
        # Auto-fill from MyDetails (Name + Role)
        if self._my_details.get("full_name"):
            author_text = self._my_details["full_name"]
            if self._my_details.get("role_title"):
                author_text += f" ({self._my_details['role_title']})"
            self.popup_author_original.setText(author_text)
        layout.addWidget(self.popup_author_original)

        # Author of update report
        author_update_lbl = QLabel("Author of Update Report:")
        author_update_lbl.setStyleSheet(label_style)
        layout.addWidget(author_update_lbl)
        self.popup_author_update = QLineEdit()
        self.popup_author_update.setStyleSheet(input_style)
        layout.addWidget(self.popup_author_update)

        # Supervisor
        supervisor_lbl = QLabel("Supervisor:")
        supervisor_lbl.setStyleSheet(label_style)
        layout.addWidget(supervisor_lbl)
        self.popup_supervisor = QLineEdit()
        self.popup_supervisor.setStyleSheet(input_style)
        layout.addWidget(self.popup_supervisor)

        # Report sent for review to
        review_lbl = QLabel("Report Sent for Review To:")
        review_lbl.setStyleSheet(label_style)
        layout.addWidget(review_lbl)
        self.popup_review_to = QLineEdit()
        self.popup_review_to.setStyleSheet(input_style)
        layout.addWidget(self.popup_review_to)

        # Date of original report
        date_orig_lbl = QLabel("Date of Original Report:")
        date_orig_lbl.setStyleSheet(label_style)
        layout.addWidget(date_orig_lbl)
        self.popup_date_original = QDateEdit()
        self.popup_date_original.setCalendarPopup(True)
        self.popup_date_original.setDisplayFormat("MMMM yyyy")
        self.popup_date_original.setStyleSheet(input_style)
        layout.addWidget(self.popup_date_original)

        # Date of update report
        date_update_lbl = QLabel("Date of Update Report:")
        date_update_lbl.setStyleSheet(label_style)
        layout.addWidget(date_update_lbl)
        self.popup_date_update = QDateEdit()
        self.popup_date_update.setCalendarPopup(True)
        self.popup_date_update.setDisplayFormat("MMMM yyyy")
        self.popup_date_update.setDate(QDate.currentDate())
        self.popup_date_update.setStyleSheet(input_style)
        layout.addWidget(self.popup_date_update)

        # Date next update due
        date_next_lbl = QLabel("Date Next Update Due:")
        date_next_lbl.setStyleSheet(label_style)
        layout.addWidget(date_next_lbl)
        self.popup_date_next = QDateEdit()
        self.popup_date_next.setCalendarPopup(True)
        self.popup_date_next.setDisplayFormat("MMMM yyyy")
        next_date = QDate.currentDate().addMonths(6)
        self.popup_date_next.setDate(next_date)
        self.popup_date_next.setStyleSheet(input_style)
        layout.addWidget(self.popup_date_next)

        layout.addStretch()

        def generate():
            parts = []
            tools = self.popup_tools.text().strip()
            if tools:
                parts.append(f"STRUCTURED CLINICAL JUDGMENT TOOLS: {tools}")
            author_orig = self.popup_author_original.text().strip()
            if author_orig:
                parts.append(f"AUTHOR OF ORIGINAL REPORT: {author_orig}")
            author_update = self.popup_author_update.text().strip()
            if author_update:
                parts.append(f"AUTHOR OF UPDATE REPORT: {author_update}")
            supervisor = self.popup_supervisor.text().strip()
            if supervisor:
                parts.append(f"Under the Supervision of: {supervisor}")
            review = self.popup_review_to.text().strip()
            if review:
                parts.append(f"REPORT SENT FOR REVIEW TO: {review}")
            date_orig = self.popup_date_original.date().toString("MMMM yyyy")
            parts.append(f"Date of original report: {date_orig}")
            date_update = self.popup_date_update.date().toString("MMMM yyyy")
            parts.append(f"Date of update report: {date_update}")
            date_next = self.popup_date_next.date().toString("MMMM yyyy")
            parts.append(f"Date next update due: {date_next}")
            return "\n".join(parts)

        self._connect_preview_updates("assessment_details", [
            self.popup_tools, self.popup_author_original, self.popup_author_update,
            self.popup_supervisor, self.popup_review_to,
            self.popup_date_original, self.popup_date_update, self.popup_date_next
        ])
        self._add_send_button(layout, "assessment_details", generate)

    def _build_popup_sources(self):
        """Build sources of information popup."""
        container, layout = self._create_popup_container("sources")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("List all sources of information used for this assessment:")
        lbl.setStyleSheet(label_style)
        lbl.setWordWrap(True)
        layout.addWidget(lbl)

        self.popup_sources_text = QTextEdit()
        self.popup_sources_text.setPlaceholderText("• Tribunal Hearing, 15th March 2022\n• Discharge Summary\n• Care notes from...\n• Previous HCR-20 report")
        self.popup_sources_text.setStyleSheet(input_style)
        self.popup_sources_text.setMinimumHeight(144)  # Reduced by 10% (was 160)
        layout.addWidget(self.popup_sources_text)

        layout.addStretch()

        def generate():
            return self.popup_sources_text.toPlainText().strip()

        self._connect_preview_updates("sources", [self.popup_sources_text])
        self._add_send_button(layout, "sources", generate)

    def _build_hcr_item_popup(self, key: str, code: str, description: str, category: str):
        """Build popup for an HCR-20 item (H1-H10, C1-C5, R1-R5)."""
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        presence_no.setStyleSheet(radio_style)
        presence_partial.setStyleSheet(radio_style)
        presence_yes.setStyleSheet(radio_style)
        presence_omit.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addWidget(presence_no)
        presence_row.addWidget(presence_partial)
        presence_row.addWidget(presence_yes)
        presence_row.addWidget(presence_omit)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        relevance_low.setStyleSheet(radio_style)
        relevance_mod.setStyleSheet(radio_style)
        relevance_high.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addWidget(relevance_low)
        relevance_row.addWidget(relevance_mod)
        relevance_row.addWidget(relevance_high)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        # Narrative/Evidence
        evidence_lbl = QLabel("Evidence/Narrative:")
        evidence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(evidence_lbl)

        evidence_text = QTextEdit()
        evidence_text.setPlaceholderText("Describe the evidence and reasoning for this scoring...")
        evidence_text.setStyleSheet(input_style)
        evidence_text.setMinimumHeight(108)  # Reduced by 10% (was 120)
        layout.addWidget(evidence_text)

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_evidence", evidence_text)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            # Presence
            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            # Relevance
            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            # Evidence
            evidence = evidence_text.toPlainText().strip()
            if evidence:
                parts.append(f"\n{evidence}")

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high, evidence_text
        ])
        self._add_send_button(layout, key, generate)

    def _build_hcr_item_popup_with_subsections(self, key: str, code: str, description: str, subsections: list):
        """Build popup for an HCR-20 item with custom subsections.

        Args:
            key: The popup key (e.g., 'h1', 'c1')
            code: The item code (e.g., 'H1', 'C1')
            description: The item description
            subsections: List of (label, placeholder) tuples for each subsection
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"
        section_label_style = "font-size: 16px; font-weight: 600; color: #059669; margin-top: 8px;"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        presence_no.setStyleSheet(radio_style)
        presence_partial.setStyleSheet(radio_style)
        presence_yes.setStyleSheet(radio_style)
        presence_omit.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addWidget(presence_no)
        presence_row.addWidget(presence_partial)
        presence_row.addWidget(presence_yes)
        presence_row.addWidget(presence_omit)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        relevance_low.setStyleSheet(radio_style)
        relevance_mod.setStyleSheet(radio_style)
        relevance_high.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addWidget(relevance_low)
        relevance_row.addWidget(relevance_mod)
        relevance_row.addWidget(relevance_high)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        # === SUBSECTIONS (Collapsible) ===
        text_widgets = []
        subsection_data = []

        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(250)
            subsections_section._min_height = 80
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)

            for label_text, placeholder in subsections:
                sub_lbl = QLabel(f"{label_text}")
                sub_lbl.setStyleSheet(section_label_style + " background: transparent; border: none;")
                subsections_content_layout.addWidget(sub_lbl)

                sub_text = QTextEdit()
                sub_text.setPlaceholderText(placeholder)
                sub_text.setStyleSheet(input_style + " background: white;")
                sub_text.setMinimumHeight(58)
                subsections_content_layout.addWidget(sub_text)
                subsections_content_layout.addWidget(DragResizeBar(sub_text))

                text_widgets.append(sub_text)
                subsection_data.append((label_text, sub_text))

                safe_name = label_text.lower().replace(" ", "_").replace("/", "_").replace(":", "").replace("–", "_").replace("-", "_")
                setattr(self, f"popup_{key}_{safe_name}", sub_text)

            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)
        else:
            for label_text, placeholder in subsections:
                sub_lbl = QLabel(f"{label_text}")
                sub_lbl.setStyleSheet(section_label_style)
                layout.addWidget(sub_lbl)

                sub_text = QTextEdit()
                sub_text.setPlaceholderText(placeholder)
                sub_text.setStyleSheet(input_style)
                sub_text.setMinimumHeight(58)
                layout.addWidget(sub_text)
                layout.addWidget(DragResizeBar(sub_text))

                text_widgets.append(sub_text)
                subsection_data.append((label_text, sub_text))

                safe_name = label_text.lower().replace(" ", "_").replace("/", "_").replace(":", "").replace("–", "_").replace("-", "_")
                setattr(self, f"popup_{key}_{safe_name}", sub_text)

        # === IMPORTED DATA (Gold/Yellow Collapsible Section with dated entries) ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(250)
            imported_section._min_height = 100
            imported_section._max_height = 500
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(8, 8, 8, 8)
            imported_content_layout.setSpacing(6)

            # Scroll area for entries
            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            imported_scroll.setStyleSheet("""
                QScrollArea { background: transparent; border: none; }
                QScrollArea > QWidget > QWidget { background: transparent; }
            """)

            # Container for entry boxes
            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(8)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            imported_section.setVisible(True)
            layout.addWidget(imported_section)

            # Store references for entries
            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])
        else:
            # Fallback if CollapsibleSection not available - use simple layout
            app_lbl = QLabel("Imported Data:")
            app_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #806000; margin-top: 16px;")
            layout.addWidget(app_lbl)

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: rgba(255, 248, 220, 0.95); border-radius: 8px;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(8, 8, 8, 8)
            imported_entries_layout.setSpacing(8)
            layout.addWidget(imported_entries_container)

            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            # Presence
            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            # Relevance
            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            # Subsection evidence
            for label_text, text_widget in subsection_data:
                evidence = text_widget.toPlainText().strip()
                if evidence:
                    parts.append(f"\n{label_text}\n{evidence}")

            # Imported Data from checked entries
            imported_checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
            imported_parts = []
            for cb in imported_checkboxes:
                if cb.isChecked():
                    full_text = cb.property("full_text")
                    if full_text:
                        imported_parts.append(full_text)
            if imported_parts:
                parts.append(f"\nImported Data:\n" + "\n\n".join(imported_parts))

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high,
        ] + text_widgets)
        self._add_send_button(layout, key, generate)

    def _build_hcr_h3_popup(self, key: str, code: str, description: str, subsections: list):
        """Build H3 (Relationships) popup with SEPARATE imported data sections for each relationship type.

        Args:
            key: The popup key ('h3')
            code: The item code ('H3')
            description: The item description
            subsections: List of (label, placeholder) tuples for Intimate and Non-intimate
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"
        section_label_style = "font-size: 16px; font-weight: 600; color: #059669; margin-top: 8px;"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        presence_no.setStyleSheet(radio_style)
        presence_partial.setStyleSheet(radio_style)
        presence_yes.setStyleSheet(radio_style)
        presence_omit.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addWidget(presence_no)
        presence_row.addWidget(presence_partial)
        presence_row.addWidget(presence_yes)
        presence_row.addWidget(presence_omit)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        relevance_low.setStyleSheet(radio_style)
        relevance_mod.setStyleSheet(radio_style)
        relevance_high.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addWidget(relevance_low)
        relevance_row.addWidget(relevance_mod)
        relevance_row.addWidget(relevance_high)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        # === SUBSECTIONS WITH THEIR OWN IMPORTED DATA SECTIONS ===
        text_widgets = []
        subsection_data = []
        subsection_keys = ["intimate", "non_intimate"]

        # === PARENT COLLAPSIBLE SECTION FOR ALL SUBSECTIONS ===
        if CollapsibleSection:
            subsections_parent = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_parent.set_content_height(400)
            subsections_parent._min_height = 150
            subsections_parent._max_height = 700

            subsections_parent.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_parent.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content_widget = QWidget()
            subsections_content_widget.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content_widget)
            subsections_content_layout.setContentsMargins(8, 8, 8, 8)
            subsections_content_layout.setSpacing(6)

            target_layout = subsections_content_layout
        else:
            target_layout = layout

        for idx, (label_text, placeholder) in enumerate(subsections):
            subsection_key = subsection_keys[idx] if idx < len(subsection_keys) else f"subsection_{idx}"

            # Subsection label
            sub_lbl = QLabel(f"{label_text}")
            sub_lbl.setStyleSheet(section_label_style)
            target_layout.addWidget(sub_lbl)

            # Subsection text field
            sub_text = QTextEdit()
            sub_text.setPlaceholderText(placeholder)
            sub_text.setStyleSheet(input_style)
            sub_text.setMinimumHeight(58)  # Reduced by 10% (was 64)
            target_layout.addWidget(sub_text)
            target_layout.addWidget(DragResizeBar(sub_text))

            text_widgets.append(sub_text)
            subsection_data.append((label_text, sub_text))

            # Store reference with sanitized name
            safe_name = label_text.lower().replace(" ", "_").replace("/", "_").replace(":", "").replace("–", "_").replace("-", "_")
            setattr(self, f"popup_{key}_{safe_name}", sub_text)

            # === IMPORTED DATA SECTION FOR THIS SUBSECTION ===
            if CollapsibleSection:
                # Determine section title based on subsection
                if "intimate" in label_text.lower() and "non" not in label_text.lower():
                    section_title = "Intimate Relationships Imported Data"
                    section_key = "intimate"
                else:
                    section_title = "Non-intimate Relationships Imported Data"
                    section_key = "non_intimate"

                imported_section = CollapsibleSection(section_title, start_collapsed=True)
                imported_section.set_content_height(200)
                imported_section._min_height = 80
                imported_section._max_height = 400

                # Use different colors for each section
                if section_key == "intimate":
                    bg_color = "rgba(255, 230, 230, 0.95)"
                    border_color = "rgba(180, 100, 100, 0.4)"
                    title_color = "#8B0000"
                else:
                    bg_color = "rgba(230, 255, 230, 0.95)"
                    border_color = "rgba(100, 180, 100, 0.4)"
                    title_color = "#006400"

                imported_section.set_header_style(f"""
                    QFrame {{
                        background: {bg_color};
                        border: 1px solid {border_color};
                        border-radius: 6px 6px 0 0;
                    }}
                """)
                imported_section.set_title_style(f"""
                    QLabel {{
                        font-size: 16px;
                        font-weight: 600;
                        color: {title_color};
                        background: transparent;
                        border: none;
                    }}
                """)

                imported_content = QWidget()
                imported_content.setStyleSheet(f"""
                    QWidget {{
                        background: {bg_color};
                        border: 1px solid {border_color};
                        border-top: none;
                        border-radius: 0 0 12px 12px;
                    }}
                    QCheckBox {{
                        background: transparent;
                        border: none;
                    }}
                """)

                imported_content_layout = QVBoxLayout(imported_content)
                imported_content_layout.setContentsMargins(8, 8, 8, 8)
                imported_content_layout.setSpacing(6)

                # Scroll area for entries
                imported_scroll = QScrollArea()
                imported_scroll.setWidgetResizable(True)
                imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
                imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                imported_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                imported_scroll.setStyleSheet("""
                    QScrollArea { background: transparent; border: none; }
                    QScrollArea > QWidget > QWidget { background: transparent; }
                """)

                # Container for entry boxes
                imported_entries_container = QWidget()
                imported_entries_container.setStyleSheet("background: transparent;")
                imported_entries_layout = QVBoxLayout(imported_entries_container)
                imported_entries_layout.setContentsMargins(2, 2, 2, 2)
                imported_entries_layout.setSpacing(8)
                imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

                imported_scroll.setWidget(imported_entries_container)
                imported_content_layout.addWidget(imported_scroll)

                imported_section.set_content(imported_content)
                imported_section.setVisible(True)
                target_layout.addWidget(imported_section)

                # Store references for this subsection's imported data
                setattr(self, f"popup_{key}_{section_key}_imported_section", imported_section)
                setattr(self, f"popup_{key}_{section_key}_imported_entries_layout", imported_entries_layout)
                setattr(self, f"popup_{key}_{section_key}_imported_checkboxes", [])
                setattr(self, f"popup_{key}_{section_key}_text_widget", sub_text)

        # Finalize and add parent CollapsibleSection to main layout
        if CollapsibleSection:
            subsections_parent.set_content(subsections_content_widget)
            layout.addWidget(subsections_parent)

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            # Presence
            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            # Relevance
            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            # Subsection evidence (includes checked imported data)
            for label_text, text_widget in subsection_data:
                evidence = text_widget.toPlainText().strip()
                if evidence:
                    parts.append(f"\n{label_text}\n{evidence}")

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high,
        ] + text_widgets)
        self._add_send_button(layout, key, generate)

    def _build_hcr_h4_popup(self, key: str, code: str, description: str, subsections: list):
        """Build H4 (Employment) popup with SEPARATE imported data sections for Education and Employment.

        Args:
            key: The popup key ('h4')
            code: The item code ('H4')
            description: The item description
            subsections: List of (label, placeholder) tuples for Education and Employment
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"
        section_label_style = "font-size: 16px; font-weight: 600; color: #059669; margin-top: 8px;"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        presence_no.setStyleSheet(radio_style)
        presence_partial.setStyleSheet(radio_style)
        presence_yes.setStyleSheet(radio_style)
        presence_omit.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addWidget(presence_no)
        presence_row.addWidget(presence_partial)
        presence_row.addWidget(presence_yes)
        presence_row.addWidget(presence_omit)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        relevance_low.setStyleSheet(radio_style)
        relevance_mod.setStyleSheet(radio_style)
        relevance_high.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addWidget(relevance_low)
        relevance_row.addWidget(relevance_mod)
        relevance_row.addWidget(relevance_high)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        # === SUBSECTIONS WITH THEIR OWN IMPORTED DATA SECTIONS ===
        text_widgets = []
        subsection_data = []
        subsection_keys = ["education", "employment"]

        # === PARENT COLLAPSIBLE SECTION FOR ALL SUBSECTIONS ===
        if CollapsibleSection:
            subsections_parent = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_parent.set_content_height(400)
            subsections_parent._min_height = 150
            subsections_parent._max_height = 700

            subsections_parent.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_parent.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content_widget = QWidget()
            subsections_content_widget.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content_widget)
            subsections_content_layout.setContentsMargins(8, 8, 8, 8)
            subsections_content_layout.setSpacing(6)

            target_layout = subsections_content_layout
        else:
            target_layout = layout

        for idx, (label_text, placeholder) in enumerate(subsections):
            subsection_key = subsection_keys[idx] if idx < len(subsection_keys) else f"subsection_{idx}"

            # Subsection label
            sub_lbl = QLabel(f"{label_text}")
            sub_lbl.setStyleSheet(section_label_style)
            target_layout.addWidget(sub_lbl)

            # Subsection text field
            sub_text = QTextEdit()
            sub_text.setPlaceholderText(placeholder)
            sub_text.setStyleSheet(input_style)
            sub_text.setMinimumHeight(58)  # Reduced by 10% (was 64)
            target_layout.addWidget(sub_text)
            target_layout.addWidget(DragResizeBar(sub_text))

            text_widgets.append(sub_text)
            subsection_data.append((label_text, sub_text))

            # Store reference with sanitized name
            safe_name = label_text.lower().replace(" ", "_").replace("/", "_").replace(":", "").replace("–", "_").replace("-", "_")
            setattr(self, f"popup_{key}_{safe_name}", sub_text)

            # === IMPORTED DATA SECTION FOR THIS SUBSECTION ===
            if CollapsibleSection:
                # Determine section title based on subsection
                if "education" in label_text.lower():
                    section_title = "Education Imported Data"
                    section_key = "education"
                    bg_color = "rgba(230, 230, 255, 0.95)"
                    border_color = "rgba(100, 100, 180, 0.4)"
                    title_color = "#00008B"
                else:
                    section_title = "Employment Imported Data"
                    section_key = "employment"
                    bg_color = "rgba(255, 245, 230, 0.95)"
                    border_color = "rgba(180, 140, 100, 0.4)"
                    title_color = "#8B4500"

                imported_section = CollapsibleSection(section_title, start_collapsed=True)
                imported_section.set_content_height(200)
                imported_section._min_height = 80
                imported_section._max_height = 400

                imported_section.set_header_style(f"""
                    QFrame {{
                        background: {bg_color};
                        border: 1px solid {border_color};
                        border-radius: 6px 6px 0 0;
                    }}
                """)
                imported_section.set_title_style(f"""
                    QLabel {{
                        font-size: 16px;
                        font-weight: 600;
                        color: {title_color};
                        background: transparent;
                        border: none;
                    }}
                """)

                imported_content = QWidget()
                imported_content.setStyleSheet(f"""
                    QWidget {{
                        background: {bg_color};
                        border: 1px solid {border_color};
                        border-top: none;
                        border-radius: 0 0 12px 12px;
                    }}
                    QCheckBox {{
                        background: transparent;
                        border: none;
                    }}
                """)

                imported_content_layout = QVBoxLayout(imported_content)
                imported_content_layout.setContentsMargins(8, 8, 8, 8)
                imported_content_layout.setSpacing(6)

                # Scroll area for entries
                imported_scroll = QScrollArea()
                imported_scroll.setWidgetResizable(True)
                imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
                imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
                imported_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
                imported_scroll.setStyleSheet("""
                    QScrollArea { background: transparent; border: none; }
                    QScrollArea > QWidget > QWidget { background: transparent; }
                """)

                # Container for entry boxes
                imported_entries_container = QWidget()
                imported_entries_container.setStyleSheet("background: transparent;")
                imported_entries_layout = QVBoxLayout(imported_entries_container)
                imported_entries_layout.setContentsMargins(2, 2, 2, 2)
                imported_entries_layout.setSpacing(8)
                imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

                imported_scroll.setWidget(imported_entries_container)
                imported_content_layout.addWidget(imported_scroll)

                imported_section.set_content(imported_content)
                imported_section.setVisible(True)
                target_layout.addWidget(imported_section)

                # Store references for this subsection's imported data
                setattr(self, f"popup_{key}_{section_key}_imported_section", imported_section)
                setattr(self, f"popup_{key}_{section_key}_imported_entries_layout", imported_entries_layout)
                setattr(self, f"popup_{key}_{section_key}_imported_checkboxes", [])
                setattr(self, f"popup_{key}_{section_key}_text_widget", sub_text)

        # Finalize and add parent CollapsibleSection to main layout
        if CollapsibleSection:
            subsections_parent.set_content(subsections_content_widget)
            layout.addWidget(subsections_parent)

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            # Presence
            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            # Relevance
            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            # Subsection evidence (includes checked imported data)
            for label_text, text_widget in subsection_data:
                evidence = text_widget.toPlainText().strip()
                if evidence:
                    parts.append(f"\n{label_text}\n{evidence}")

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high,
        ] + text_widgets)
        self._add_send_button(layout, key, generate)

    def _build_hcr_h7_popup(self, key: str, code: str, description: str, subsections: list):
        """Build H7 (Personality Disorder) popup with PD type buttons and trait checkboxes.

        Based on HCR-20V3 manual guidance for H7 - personality disorder or
        personality dysfunction relevant to violence risk.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS (Collapsible) ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === PERSONALITY DISORDER FEATURES CONTAINER ===
        pd_container = QFrame()
        pd_container.setStyleSheet("""
            QFrame {
                background: rgba(237, 233, 254, 0.5);
                border: 2px solid rgba(124, 58, 237, 0.4);
                border-radius: 10px;
            }
        """)
        pd_layout = QVBoxLayout(pd_container)
        pd_layout.setContentsMargins(12, 10, 12, 10)
        pd_layout.setSpacing(8)

        pd_lbl = QLabel("Personality Disorder Features")
        pd_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #6d28d9; background: transparent; border: none;")
        pd_layout.addWidget(pd_lbl)

        pd_text = QTextEdit()
        pd_text.setPlaceholderText("Document personality disorder diagnosis, features, and traits...")
        pd_text.setStyleSheet(input_style + " background: white;")
        pd_text.setMinimumHeight(58)
        pd_layout.addWidget(pd_text)
        pd_layout.addWidget(DragResizeBar(pd_text))
        text_widgets.append(pd_text)
        subsection_data.append(("Personality Disorder Features", pd_text))
        setattr(self, "popup_h7_pd_features", pd_text)

        # PD Type buttons in grid layout (2 rows)
        pd_buttons_frame = QFrame()
        pd_buttons_frame.setStyleSheet("background: transparent; border: none;")
        pd_buttons_layout = QGridLayout(pd_buttons_frame)
        pd_buttons_layout.setContentsMargins(0, 4, 0, 4)
        pd_buttons_layout.setSpacing(4)

        # Define PD types with their colors (ICD-10 aligned) - no Narcissistic, expanded Other PDs
        PD_TYPES = [
            # Row 1
            ("Dissocial", "Dissocial (F60.2)", "#dc2626"),
            ("EUPD-B", "EUPD Borderline (F60.31)", "#ea580c"),
            ("EUPD-I", "EUPD Impulsive (F60.30)", "#f59e0b"),
            ("Paranoid", "Paranoid (F60.0)", "#15803d"),
            ("Schizoid", "Schizoid (F60.1)", "#0891b2"),
            # Row 2
            ("Histrionic", "Histrionic (F60.4)", "#c026d3"),
            ("Anankastic", "Anankastic (F60.5)", "#4f46e5"),
            ("Anxious", "Anxious (F60.6)", "#0d9488"),
            ("Dependent", "Dependent (F60.7)", "#6366f1"),
        ]

        # Stacked widget for trait panels
        pd_stack = QStackedWidget()
        pd_stack.setStyleSheet("background: transparent;")

        pd_buttons = []
        all_pd_checkboxes = []

        # Define traits for each PD type (ICD-10 criteria)
        PD_TRAITS = {
            "Dissocial": [
                ("dissocial_unconcern", "Callous unconcern for feelings of others"),
                ("dissocial_irresponsibility", "Gross and persistent irresponsibility"),
                ("dissocial_incapacity_relations", "Incapacity to maintain enduring relationships"),
                ("dissocial_low_frustration", "Very low tolerance to frustration"),
                ("dissocial_aggression", "Low threshold for discharge of aggression"),
                ("dissocial_incapacity_guilt", "Incapacity to experience guilt"),
                ("dissocial_blame_others", "Marked proneness to blame others"),
                ("dissocial_rationalise", "Plausible rationalisation for behaviour"),
            ],
            "EUPD-B": [
                ("eupd_b_abandonment", "Frantic efforts to avoid abandonment"),
                ("eupd_b_unstable_relations", "Unstable and intense relationships"),
                ("eupd_b_identity", "Identity disturbance"),
                ("eupd_b_impulsivity", "Impulsivity in potentially damaging areas"),
                ("eupd_b_suicidal", "Recurrent suicidal/self-harm behaviour"),
                ("eupd_b_affective", "Affective instability"),
                ("eupd_b_emptiness", "Chronic feelings of emptiness"),
                ("eupd_b_anger", "Inappropriate, intense anger"),
                ("eupd_b_dissociation", "Transient paranoia or dissociation"),
            ],
            "EUPD-I": [
                ("eupd_i_act_unexpectedly", "Acts unexpectedly without considering consequences"),
                ("eupd_i_quarrelsome", "Tendency to quarrelsome behaviour and conflicts"),
                ("eupd_i_anger_outbursts", "Liability to outbursts of anger or violence"),
                ("eupd_i_no_persistence", "Difficulty maintaining actions without immediate reward"),
                ("eupd_i_unstable_mood", "Unstable and capricious mood"),
            ],
            "Paranoid": [
                ("paranoid_suspects", "Suspects others are exploiting or harming"),
                ("paranoid_doubts_loyalty", "Preoccupied with doubts about loyalty"),
                ("paranoid_reluctant_confide", "Reluctant to confide in others"),
                ("paranoid_reads_threats", "Reads hidden demeaning meanings"),
                ("paranoid_bears_grudges", "Persistently bears grudges"),
                ("paranoid_perceives_attacks", "Perceives attacks on character"),
                ("paranoid_suspicious_fidelity", "Recurrent suspicions about fidelity"),
            ],
            "Schizoid": [
                ("schizoid_no_pleasure", "Few activities give pleasure"),
                ("schizoid_cold", "Emotional coldness, detachment, flat affect"),
                ("schizoid_limited_warmth", "Limited capacity to express warmth or anger"),
                ("schizoid_indifferent", "Apparent indifference to praise or criticism"),
                ("schizoid_little_interest_sex", "Little interest in sexual experiences"),
                ("schizoid_solitary", "Preference for solitary activities"),
                ("schizoid_fantasy", "Excessive preoccupation with fantasy and introspection"),
                ("schizoid_no_confidants", "No close friends or confiding relationships"),
                ("schizoid_insensitive", "Insensitivity to social norms and conventions"),
            ],
            "Histrionic": [
                ("histrionic_attention", "Discomfort when not centre of attention"),
                ("histrionic_seductive", "Inappropriately seductive or provocative"),
                ("histrionic_shallow_emotion", "Rapidly shifting and shallow emotions"),
                ("histrionic_appearance", "Uses appearance to draw attention"),
                ("histrionic_impressionistic", "Impressionistic speech lacking detail"),
                ("histrionic_dramatic", "Self-dramatisation, theatricality, exaggerated emotion"),
                ("histrionic_suggestible", "Easily influenced by others or circumstances"),
                ("histrionic_intimacy", "Considers relationships more intimate than they are"),
            ],
            "Anankastic": [
                ("anankastic_doubt", "Excessive doubt and caution"),
                ("anankastic_detail", "Preoccupation with details, rules, lists, order"),
                ("anankastic_perfectionism", "Perfectionism that interferes with completion"),
                ("anankastic_conscientious", "Excessive conscientiousness and scrupulousness"),
                ("anankastic_pleasure", "Preoccupation with productivity to exclusion of pleasure"),
                ("anankastic_pedantic", "Excessive pedantry and adherence to convention"),
                ("anankastic_rigid", "Rigidity and stubbornness"),
                ("anankastic_insistence", "Unreasonable insistence others do things their way"),
            ],
            "Anxious": [
                ("anxious_tension", "Persistent feelings of tension and apprehension"),
                ("anxious_inferior", "Beliefs of social ineptness and inferiority"),
                ("anxious_criticism", "Excessive preoccupation with criticism or rejection"),
                ("anxious_unwilling", "Unwilling to become involved unless certain of being liked"),
                ("anxious_restricted", "Restrictions in lifestyle due to need for security"),
                ("anxious_avoids_activities", "Avoids activities involving significant interpersonal contact"),
            ],
            "Dependent": [
                ("dependent_encourage", "Encourages or allows others to make decisions"),
                ("dependent_subordinates", "Subordinates own needs to those of others"),
                ("dependent_unwilling_demands", "Unwilling to make reasonable demands on others"),
                ("dependent_helpless", "Feels uncomfortable or helpless when alone"),
                ("dependent_abandonment", "Preoccupied with fears of being left to care for self"),
                ("dependent_limited_capacity", "Limited capacity to make everyday decisions without advice"),
            ],
        }

        for idx, (pd_abbrev, pd_name, pd_color) in enumerate(PD_TYPES):
            btn = QPushButton(pd_abbrev)
            btn.setFixedSize(90, 26)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.setCheckable(True)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: #f3f4f6;
                    color: #374151;
                    border: 1px solid #d1d5db;
                    border-radius: 6px;
                    font-size: 11px;
                    font-weight: 600;
                    padding: 2px 4px;
                }}
                QPushButton:hover {{
                    background: {pd_color};
                    color: white;
                    border-color: {pd_color};
                }}
                QPushButton:checked {{
                    background: {pd_color};
                    color: white;
                    border-color: {pd_color};
                }}
            """)
            btn.setToolTip(pd_name)
            row = idx // 5  # 5 buttons per row
            col = idx % 5
            pd_buttons_layout.addWidget(btn, row, col)
            pd_buttons.append((btn, pd_abbrev))

            # Create trait panel for this PD
            trait_panel = QWidget()
            trait_panel.setStyleSheet(f"""
                QWidget {{
                    background: rgba(245, 243, 255, 0.95);
                    border: 1px solid rgba(124, 58, 237, 0.2);
                    border-radius: 8px;
                }}
                QCheckBox {{
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }}
            """)
            trait_layout = QVBoxLayout(trait_panel)
            trait_layout.setContentsMargins(10, 8, 10, 8)
            trait_layout.setSpacing(3)

            panel_title = QLabel(f"{pd_name} Traits")
            panel_title.setStyleSheet(f"font-size: 13px; font-weight: 700; color: {pd_color}; background: transparent;")
            trait_layout.addWidget(panel_title)

            panel_checkboxes = []
            for trait_key, trait_label in PD_TRAITS.get(pd_abbrev, []):
                cb = QCheckBox(trait_label)
                cb.setProperty("pd_type", pd_abbrev)
                cb.setProperty("trait_key", trait_key)
                cb.stateChanged.connect(self._update_h7_pd_narrative)
                trait_layout.addWidget(cb)
                panel_checkboxes.append(cb)
                all_pd_checkboxes.append(cb)

            trait_layout.addStretch()
            pd_stack.addWidget(trait_panel)

        pd_layout.addWidget(pd_buttons_frame)

        # Connect buttons to show corresponding panel
        def make_button_handler(idx, btn):
            def handler(checked):
                if checked:
                    # Uncheck other buttons
                    for other_btn, _ in pd_buttons:
                        if other_btn != btn:
                            other_btn.setChecked(False)
                    pd_stack.setCurrentIndex(idx)
                    pd_stack.setVisible(True)
                else:
                    # If unchecking, hide if no button is checked
                    any_checked = any(b.isChecked() for b, _ in pd_buttons)
                    if not any_checked:
                        pd_stack.setVisible(False)
            return handler

        for idx, (btn, _) in enumerate(pd_buttons):
            btn.clicked.connect(make_button_handler(idx, btn))

        # Initially hide the stack
        pd_stack.setVisible(False)
        pd_stack.setMinimumHeight(180)
        pd_stack.setMaximumHeight(280)
        pd_layout.addWidget(pd_stack)

        setattr(self, "popup_h7_pd_checkboxes", all_pd_checkboxes)
        setattr(self, "popup_h7_pd_buttons", pd_buttons)

        subsections_target_layout.addWidget(pd_container)

        # === IMPACT ON FUNCTIONING CONTAINER ===
        impact_container = QFrame()
        impact_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(180, 83, 9, 0.4);
                border-radius: 10px;
            }
        """)
        impact_layout = QVBoxLayout(impact_container)
        impact_layout.setContentsMargins(12, 10, 12, 10)
        impact_layout.setSpacing(8)

        impact_lbl = QLabel("Impact on Functioning")
        impact_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        impact_layout.addWidget(impact_lbl)

        impact_text = QTextEdit()
        impact_text.setPlaceholderText("Document how personality difficulties impact relationships, employment, treatment, and risk...")
        impact_text.setStyleSheet(input_style + " background: white;")
        impact_text.setMinimumHeight(58)
        impact_layout.addWidget(impact_text)
        impact_layout.addWidget(DragResizeBar(impact_text))
        text_widgets.append(impact_text)
        subsection_data.append(("Impact on Functioning", impact_text))
        setattr(self, "popup_h7_impact_functioning", impact_text)

        if CollapsibleSection:
            impact_section = CollapsibleSection("Select Applicable Impacts", start_collapsed=False)
            impact_section.set_content_height(180)
            impact_section._min_height = 80
            impact_section._max_height = 280

            impact_section.set_header_style("""
                QFrame {
                    background: rgba(180, 83, 9, 0.15);
                    border: 1px solid rgba(180, 83, 9, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            impact_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #92400e;
                    background: transparent;
                    border: none;
                }
            """)

            impact_content = QWidget()
            impact_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 251, 235, 0.95);
                    border: 1px solid rgba(180, 83, 9, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            impact_cb_layout = QVBoxLayout(impact_content)
            impact_cb_layout.setContentsMargins(10, 6, 10, 6)
            impact_cb_layout.setSpacing(3)

            H7_IMPACT = [
                ("Relationships", "#b45309", [
                    ("impact_intimate", "Difficulties maintaining intimate relationships"),
                    ("impact_family", "Conflictual or estranged family relationships"),
                    ("impact_social", "Poor social relationships / isolation"),
                    ("impact_professional", "Difficulties with professional relationships"),
                ]),
                ("Employment/Education", "#d97706", [
                    ("impact_job_loss", "Repeated job losses due to behaviour"),
                    ("impact_work_conflict", "Frequent workplace conflicts"),
                    ("impact_underachievement", "Significant underachievement"),
                ]),
                ("Treatment/Supervision", "#f59e0b", [
                    ("impact_poor_engagement", "Poor treatment engagement"),
                    ("impact_staff_conflict", "Conflicts with clinical staff"),
                    ("impact_non_compliance", "Non-compliance with supervision"),
                    ("impact_manipulation", "Manipulative behaviour in treatment"),
                ]),
                ("Violence Risk", "#dc2626", [
                    ("impact_aggression_pattern", "Pattern of aggressive behaviour"),
                    ("impact_instrumental", "Instrumental/planned violence"),
                    ("impact_reactive", "Reactive/impulsive violence"),
                    ("impact_victim_targeting", "Targeting of specific victim types"),
                ]),
            ]

            h7_impact_checkboxes = []
            for cat_name, cat_color, items in H7_IMPACT:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                impact_cb_layout.addWidget(cat_lbl)
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("impact_key", item_key)
                    cb.stateChanged.connect(self._update_h7_pd_narrative)
                    impact_cb_layout.addWidget(cb)
                    h7_impact_checkboxes.append(cb)

            impact_section.set_content(impact_content)
            impact_layout.addWidget(impact_section)
            setattr(self, "popup_h7_impact_checkboxes", h7_impact_checkboxes)

        subsections_target_layout.addWidget(impact_container)

        # Finalize the Subsections CollapsibleSection
        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            for label_text, text_widget in subsection_data:
                evidence = text_widget.toPlainText().strip()
                if evidence:
                    parts.append(f"\n{label_text}\n{evidence}")

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high,
        ] + text_widgets)
        self._add_send_button(layout, key, generate)

    def _update_h7_pd_narrative(self):
        """Update H7 text fields with gender-sensitive narrative based on checked PD traits."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        pd_text = getattr(self, "popup_h7_pd_features", None)
        impact_text = getattr(self, "popup_h7_impact_functioning", None)
        pd_cbs = getattr(self, "popup_h7_pd_checkboxes", [])
        impact_cbs = getattr(self, "popup_h7_impact_checkboxes", [])

        # Group checked traits by PD type
        pd_groups = {}
        for cb in pd_cbs:
            if cb.isChecked():
                pd_type = cb.property("pd_type")
                trait_key = cb.property("trait_key")
                if pd_type not in pd_groups:
                    pd_groups[pd_type] = []
                pd_groups[pd_type].append(trait_key)

        # Build narrative for PD features (ICD-10 aligned)
        pd_narratives = {
            # Dissocial PD (F60.2) traits
            "dissocial_unconcern": "callous unconcern for the feelings of others",
            "dissocial_irresponsibility": "gross and persistent irresponsibility",
            "dissocial_incapacity_relations": "incapacity to maintain enduring relationships",
            "dissocial_low_frustration": "very low tolerance to frustration",
            "dissocial_aggression": "a low threshold for discharge of aggression",
            "dissocial_incapacity_guilt": "incapacity to experience guilt",
            "dissocial_blame_others": "marked proneness to blame others",
            "dissocial_rationalise": "plausible rationalisation of behaviour",
            # EUPD Borderline (F60.31) traits
            "eupd_b_abandonment": "frantic efforts to avoid abandonment",
            "eupd_b_unstable_relations": "a pattern of unstable and intense relationships",
            "eupd_b_identity": "identity disturbance",
            "eupd_b_impulsivity": "impulsivity in potentially damaging areas",
            "eupd_b_suicidal": "recurrent suicidal or self-harming behaviour",
            "eupd_b_affective": "affective instability",
            "eupd_b_emptiness": "chronic feelings of emptiness",
            "eupd_b_anger": "inappropriate and intense anger",
            "eupd_b_dissociation": "transient paranoid ideation or dissociation",
            # EUPD Impulsive (F60.30) traits
            "eupd_i_act_unexpectedly": "a tendency to act unexpectedly without considering consequences",
            "eupd_i_quarrelsome": "a tendency to quarrelsome behaviour and conflicts with others",
            "eupd_i_anger_outbursts": "liability to outbursts of anger or violence",
            "eupd_i_no_persistence": "difficulty maintaining actions that offer no immediate reward",
            "eupd_i_unstable_mood": "unstable and capricious mood",
            # Paranoid PD (F60.0) traits
            "paranoid_suspects": "suspicion that others are exploiting or harming {obj}",
            "paranoid_doubts_loyalty": "preoccupation with doubts about loyalty of others",
            "paranoid_reluctant_confide": "reluctance to confide in others",
            "paranoid_reads_threats": "tendency to read hidden threatening meanings into benign events",
            "paranoid_bears_grudges": "persistent bearing of grudges",
            "paranoid_perceives_attacks": "perception of attacks on {poss} character",
            "paranoid_suspicious_fidelity": "recurrent suspicions about partner's fidelity",
            # Schizoid PD (F60.1) traits
            "schizoid_no_pleasure": "few activities that give pleasure",
            "schizoid_cold": "emotional coldness, detachment, or flat affect",
            "schizoid_limited_warmth": "limited capacity to express warmth or anger toward others",
            "schizoid_indifferent": "apparent indifference to praise or criticism",
            "schizoid_little_interest_sex": "little interest in sexual experiences",
            "schizoid_solitary": "a preference for solitary activities",
            "schizoid_fantasy": "excessive preoccupation with fantasy and introspection",
            "schizoid_no_confidants": "lack of close friends or confiding relationships",
            "schizoid_insensitive": "insensitivity to prevailing social norms and conventions",
            # Histrionic PD (F60.4) traits
            "histrionic_attention": "discomfort when not the centre of attention",
            "histrionic_seductive": "inappropriately seductive or provocative behaviour",
            "histrionic_shallow_emotion": "rapidly shifting and shallow expression of emotions",
            "histrionic_appearance": "use of physical appearance to draw attention",
            "histrionic_impressionistic": "impressionistic speech lacking in detail",
            "histrionic_dramatic": "self-dramatisation, theatricality, and exaggerated emotion",
            "histrionic_suggestible": "easily influenced by others or circumstances",
            "histrionic_intimacy": "considering relationships more intimate than they are",
            # Anankastic PD (F60.5) traits
            "anankastic_doubt": "excessive doubt and caution",
            "anankastic_detail": "preoccupation with details, rules, lists, order, and schedules",
            "anankastic_perfectionism": "perfectionism that interferes with task completion",
            "anankastic_conscientious": "excessive conscientiousness and scrupulousness",
            "anankastic_pleasure": "preoccupation with productivity to the exclusion of pleasure",
            "anankastic_pedantic": "excessive pedantry and adherence to social conventions",
            "anankastic_rigid": "rigidity and stubbornness",
            "anankastic_insistence": "unreasonable insistence that others do things {poss} way",
            # Anxious PD (F60.6) traits
            "anxious_tension": "persistent feelings of tension and apprehension",
            "anxious_inferior": "beliefs of being socially inept or inferior to others",
            "anxious_criticism": "excessive preoccupation with criticism or rejection",
            "anxious_unwilling": "unwillingness to become involved unless certain of being liked",
            "anxious_restricted": "restrictions in lifestyle due to need for physical security",
            "anxious_avoids_activities": "avoidance of activities involving significant interpersonal contact",
            # Dependent PD (F60.7) traits
            "dependent_encourage": "encouraging or allowing others to make important decisions",
            "dependent_subordinates": "subordination of own needs to those of others",
            "dependent_unwilling_demands": "unwillingness to make reasonable demands on others",
            "dependent_helpless": "feeling uncomfortable or helpless when alone",
            "dependent_abandonment": "preoccupation with fears of being left to care for {refl}",
            "dependent_limited_capacity": "limited capacity to make everyday decisions without advice",
        }

        PD_NAMES = {
            "Dissocial": "dissocial personality disorder",
            "EUPD-B": "emotionally unstable personality disorder, borderline type",
            "EUPD-I": "emotionally unstable personality disorder, impulsive type",
            "Paranoid": "paranoid personality disorder",
            "Schizoid": "schizoid personality disorder",
            "Histrionic": "histrionic personality disorder",
            "Anankastic": "anankastic personality disorder",
            "Anxious": "anxious (avoidant) personality disorder",
            "Dependent": "dependent personality disorder",
        }

        if pd_text and pd_groups:
            narrative_parts = []
            for pd_type, traits in pd_groups.items():
                pd_name = PD_NAMES.get(pd_type, "personality disorder")
                trait_phrases = []
                for trait in traits:
                    phrase = pd_narratives.get(trait, "")
                    if phrase:
                        # Replace placeholders
                        phrase = phrase.replace("{subj}", subj).replace("{obj}", obj).replace("{poss}", poss).replace("{refl}", refl)
                        trait_phrases.append(phrase)

                if trait_phrases:
                    intro = f"{Subj} presents with features of {pd_name} including"

                    if len(trait_phrases) == 1:
                        narrative_parts.append(f"{intro} {trait_phrases[0]}.")
                    elif len(trait_phrases) == 2:
                        narrative_parts.append(f"{intro} {trait_phrases[0]} and {trait_phrases[1]}.")
                    else:
                        all_but_last = ", ".join(trait_phrases[:-1])
                        narrative_parts.append(f"{intro} {all_but_last}, and {trait_phrases[-1]}.")

            narrative = " ".join(narrative_parts)
            pd_text.setPlainText(narrative)

        # Build narrative for impact on functioning - grouped by category
        # Category phrases (without subject prefix for flowing narrative)
        relationship_phrases = {
            "impact_intimate": "difficulty maintaining intimate relationships",
            "impact_family": "conflictual or estranged family relationships",
            "impact_social": "poor social relationships and a tendency toward isolation",
            "impact_professional": "difficulties with professional relationships",
        }
        employment_phrases = {
            "impact_job_loss": "repeated job losses due to behaviour",
            "impact_work_conflict": "frequent workplace conflicts",
            "impact_underachievement": "significant underachievement",
        }
        treatment_phrases = {
            "impact_poor_engagement": "poor treatment engagement",
            "impact_staff_conflict": "conflicts with clinical staff",
            "impact_non_compliance": "non-compliance with supervision requirements",
            "impact_manipulation": "manipulative behaviour in treatment settings",
        }
        violence_phrases = {
            "impact_aggression_pattern": "a pattern of aggressive behaviour",
            "impact_instrumental": "instrumental or planned violence",
            "impact_reactive": "reactive or impulsive violence",
            "impact_victim_targeting": "targeting of specific victim types",
        }

        # Collect checked items by category
        relationship_items = []
        employment_items = []
        treatment_items = []
        violence_items = []

        for cb in impact_cbs:
            if cb.isChecked():
                key = cb.property("impact_key")
                if key in relationship_phrases:
                    relationship_items.append(relationship_phrases[key])
                elif key in employment_phrases:
                    employment_items.append(employment_phrases[key])
                elif key in treatment_phrases:
                    treatment_items.append(treatment_phrases[key])
                elif key in violence_phrases:
                    violence_items.append(violence_phrases[key])

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return items[0] + " and " + items[1]
            else:
                return ", ".join(items[:-1]) + ", and " + items[-1]

        sentences = []

        # Relationships sentence
        if relationship_items:
            sentences.append(f"{Poss} personality impacts {poss} relationships in that {subj} has {join_items(relationship_items)}.")

        # Employment sentence
        if employment_items:
            sentences.append(f"With respect to employment, {subj} has experienced {join_items(employment_items)}.")

        # Treatment/Supervision sentence
        if treatment_items:
            if sentences:
                sentences.append(f"Treatment and supervision are also affected by {poss} personality issues, with {join_items(treatment_items)}.")
            else:
                sentences.append(f"Treatment and supervision are affected by {poss} personality issues, with {join_items(treatment_items)}.")

        # Violence sentence
        if violence_items:
            if sentences:
                sentences.append(f"There is a concerning pattern of violence secondary to {poss} personality concerns, with {join_items(violence_items)}.")
            else:
                sentences.append(f"There is a concerning pattern of violence related to {poss} personality, with {join_items(violence_items)}.")

        if impact_text:
            if sentences:
                impact_text.setPlainText(" ".join(sentences))
            else:
                impact_text.clear()

    def _build_hcr_h8_popup(self, key: str, code: str, description: str, subsections: list):
        """Build H8 (Traumatic Experiences) popup with operationalized trauma categories.

        Based on HCR-20V3 manual guidance for H8 - significant traumatic experiences
        across the lifespan that are plausibly relevant to violence risk.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600

            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(8, 8, 8, 8)
            subsections_content_layout.setSpacing(8)

        # === CHILDHOOD & DEVELOPMENTAL TRAUMA CONTAINER ===
        childhood_container = QFrame()
        childhood_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(185, 28, 28, 0.4);
                border-radius: 10px;
            }
        """)
        childhood_layout = QVBoxLayout(childhood_container)
        childhood_layout.setContentsMargins(12, 10, 12, 10)
        childhood_layout.setSpacing(8)

        childhood_lbl = QLabel("Childhood & Developmental Trauma")
        childhood_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        childhood_layout.addWidget(childhood_lbl)

        childhood_text = QTextEdit()
        childhood_text.setPlaceholderText("Document abuse, neglect, adverse childhood experiences...")
        childhood_text.setStyleSheet(input_style + " background: white;")
        childhood_text.setMinimumHeight(58)
        childhood_layout.addWidget(childhood_text)
        childhood_layout.addWidget(DragResizeBar(childhood_text))
        text_widgets.append(childhood_text)
        subsection_data.append(("Childhood Trauma", childhood_text))
        setattr(self, "popup_h8_childhood_trauma", childhood_text)

        if CollapsibleSection:
            childhood_section = CollapsibleSection("Select Applicable Categories", start_collapsed=False)
            childhood_section.set_content_height(220)
            childhood_section._min_height = 80
            childhood_section._max_height = 350

            childhood_section.set_header_style("""
                QFrame {
                    background: rgba(185, 28, 28, 0.15);
                    border: 1px solid rgba(185, 28, 28, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            childhood_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            childhood_content = QWidget()
            childhood_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(185, 28, 28, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            childhood_cb_layout = QVBoxLayout(childhood_content)
            childhood_cb_layout.setContentsMargins(10, 6, 10, 6)
            childhood_cb_layout.setSpacing(3)

            H8_CHILDHOOD_TRAUMA = [
                ("Abuse", "#b91c1c", [
                    ("physical_abuse", "Physical abuse by caregivers"),
                    ("sexual_abuse_child", "Sexual abuse (any perpetrator)"),
                    ("emotional_abuse", "Emotional/psychological abuse"),
                    ("witnessed_dv", "Witnessed domestic violence in home"),
                ]),
                ("Neglect & Deprivation", "#dc2626", [
                    ("emotional_neglect", "Emotional neglect"),
                    ("physical_neglect", "Physical neglect"),
                    ("inconsistent_care", "Inconsistent or absent caregiving"),
                    ("institutional_care", "Institutional care / foster care instability"),
                    ("parental_abandonment", "Parental abandonment"),
                ]),
                ("Adverse Upbringing", "#ef4444", [
                    ("chaotic_household", "Chaotic household environment"),
                    ("parental_substance", "Parental substance misuse"),
                    ("parental_mental_illness", "Parental mental illness"),
                    ("criminal_caregivers", "Criminal or violent caregivers"),
                    ("placement_breakdowns", "Repeated placement breakdowns"),
                ]),
            ]

            h8_childhood_checkboxes = []
            for cat_name, cat_color, items in H8_CHILDHOOD_TRAUMA:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                childhood_cb_layout.addWidget(cat_lbl)
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("trauma_key", item_key)
                    cb.setProperty("trauma_type", "childhood")
                    cb.stateChanged.connect(self._update_h8_trauma_narrative)
                    childhood_cb_layout.addWidget(cb)
                    h8_childhood_checkboxes.append(cb)

            childhood_section.set_content(childhood_content)
            childhood_layout.addWidget(childhood_section)
            setattr(self, "popup_h8_childhood_checkboxes", h8_childhood_checkboxes)

        if CollapsibleSection:
            subsections_content_layout.addWidget(childhood_container)
        else:
            layout.addWidget(childhood_container)

        # === ADULT TRAUMA CONTAINER ===
        adult_container = QFrame()
        adult_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(180, 83, 9, 0.4);
                border-radius: 10px;
            }
        """)
        adult_layout = QVBoxLayout(adult_container)
        adult_layout.setContentsMargins(12, 10, 12, 10)
        adult_layout.setSpacing(8)

        adult_lbl = QLabel("Adult Trauma")
        adult_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        adult_layout.addWidget(adult_lbl)

        adult_text = QTextEdit()
        adult_text.setPlaceholderText("Document victimisation, institutional trauma, occupational trauma...")
        adult_text.setStyleSheet(input_style + " background: white;")
        adult_text.setMinimumHeight(58)
        adult_layout.addWidget(adult_text)
        adult_layout.addWidget(DragResizeBar(adult_text))
        text_widgets.append(adult_text)
        subsection_data.append(("Adult Trauma", adult_text))
        setattr(self, "popup_h8_adult_trauma", adult_text)

        if CollapsibleSection:
            adult_section = CollapsibleSection("Select Applicable Categories", start_collapsed=False)
            adult_section.set_content_height(200)
            adult_section._min_height = 80
            adult_section._max_height = 320

            adult_section.set_header_style("""
                QFrame {
                    background: rgba(180, 83, 9, 0.15);
                    border: 1px solid rgba(180, 83, 9, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            adult_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #92400e;
                    background: transparent;
                    border: none;
                }
            """)

            adult_content = QWidget()
            adult_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 251, 235, 0.95);
                    border: 1px solid rgba(180, 83, 9, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            adult_cb_layout = QVBoxLayout(adult_content)
            adult_cb_layout.setContentsMargins(10, 6, 10, 6)
            adult_cb_layout.setSpacing(3)

            H8_ADULT_TRAUMA = [
                ("Victimisation", "#b45309", [
                    ("adult_assault", "Assaults (street violence, domestic violence)"),
                    ("sexual_assault_adult", "Sexual assault / rape"),
                    ("robbery_violence", "Robbery with violence"),
                    ("stalking_coercion", "Stalking or coercive control"),
                ]),
                ("Institutional/Systemic", "#d97706", [
                    ("prison_violence", "Prison violence"),
                    ("segregation_isolation", "Segregation / prolonged isolation"),
                    ("hospital_victimisation", "Victimisation in hospital or care"),
                    ("bullying_harassment", "Bullying, harassment, exploitation"),
                ]),
                ("Occupational", "#f59e0b", [
                    ("occupational_violence", "Exposure to violence (security, military, etc.)"),
                    ("witnessed_death", "Witnessing serious injury or death"),
                    ("threats_to_life", "Threats to life"),
                ]),
            ]

            h8_adult_checkboxes = []
            for cat_name, cat_color, items in H8_ADULT_TRAUMA:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                adult_cb_layout.addWidget(cat_lbl)
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("trauma_key", item_key)
                    cb.setProperty("trauma_type", "adult")
                    cb.stateChanged.connect(self._update_h8_trauma_narrative)
                    adult_cb_layout.addWidget(cb)
                    h8_adult_checkboxes.append(cb)

            adult_section.set_content(adult_content)
            adult_layout.addWidget(adult_section)
            setattr(self, "popup_h8_adult_checkboxes", h8_adult_checkboxes)

        if CollapsibleSection:
            subsections_content_layout.addWidget(adult_container)
        else:
            layout.addWidget(adult_container)

        # === LOSS & CATASTROPHE CONTAINER ===
        loss_container = QFrame()
        loss_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(30, 64, 175, 0.4);
                border-radius: 10px;
            }
        """)
        loss_layout = QVBoxLayout(loss_container)
        loss_layout.setContentsMargins(12, 10, 12, 10)
        loss_layout.setSpacing(8)

        loss_lbl = QLabel("Trauma Linked to Loss or Catastrophe")
        loss_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1e40af; background: transparent; border: none;")
        loss_layout.addWidget(loss_lbl)

        loss_text = QTextEdit()
        loss_text.setPlaceholderText("Document bereavements, war/displacement, accidents, disasters...")
        loss_text.setStyleSheet(input_style + " background: white;")
        loss_text.setMinimumHeight(58)
        loss_layout.addWidget(loss_text)
        loss_layout.addWidget(DragResizeBar(loss_text))
        text_widgets.append(loss_text)
        subsection_data.append(("Loss/Catastrophe", loss_text))
        setattr(self, "popup_h8_loss_catastrophe", loss_text)

        if CollapsibleSection:
            loss_section = CollapsibleSection("Select Applicable Categories", start_collapsed=True)
            loss_section.set_content_height(150)
            loss_section._min_height = 60
            loss_section._max_height = 250

            loss_section.set_header_style("""
                QFrame {
                    background: rgba(30, 64, 175, 0.15);
                    border: 1px solid rgba(30, 64, 175, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            loss_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #1e40af;
                    background: transparent;
                    border: none;
                }
            """)

            loss_content = QWidget()
            loss_content.setStyleSheet("""
                QWidget {
                    background: rgba(239, 246, 255, 0.95);
                    border: 1px solid rgba(30, 64, 175, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            loss_cb_layout = QVBoxLayout(loss_content)
            loss_cb_layout.setContentsMargins(10, 6, 10, 6)
            loss_cb_layout.setSpacing(3)

            H8_LOSS_CATASTROPHE = [
                ("Bereavement", "#1e40af", [
                    ("violent_death", "Sudden or violent death of close others"),
                    ("multiple_bereavements", "Multiple bereavements"),
                ]),
                ("Displacement/Torture", "#3730a3", [
                    ("war_displacement", "War, displacement, torture"),
                    ("forced_migration", "Forced migration / asylum-related trauma"),
                ]),
                ("Accidents/Disasters", "#4f46e5", [
                    ("serious_accidents", "Serious accidents (RTA, workplace)"),
                    ("disasters", "Natural or man-made disasters"),
                ]),
            ]

            h8_loss_checkboxes = []
            for cat_name, cat_color, items in H8_LOSS_CATASTROPHE:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                loss_cb_layout.addWidget(cat_lbl)
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("trauma_key", item_key)
                    cb.setProperty("trauma_type", "loss")
                    cb.stateChanged.connect(self._update_h8_trauma_narrative)
                    loss_cb_layout.addWidget(cb)
                    h8_loss_checkboxes.append(cb)

            loss_section.set_content(loss_content)
            loss_layout.addWidget(loss_section)
            setattr(self, "popup_h8_loss_checkboxes", h8_loss_checkboxes)

        if CollapsibleSection:
            subsections_content_layout.addWidget(loss_container)
        else:
            layout.addWidget(loss_container)

        # === PSYCHOLOGICAL SEQUELAE CONTAINER ===
        sequelae_container = QFrame()
        sequelae_container.setStyleSheet("""
            QFrame {
                background: rgba(237, 233, 254, 0.5);
                border: 2px solid rgba(124, 58, 237, 0.4);
                border-radius: 10px;
            }
        """)
        sequelae_layout = QVBoxLayout(sequelae_container)
        sequelae_layout.setContentsMargins(12, 10, 12, 10)
        sequelae_layout.setSpacing(8)

        sequelae_lbl = QLabel("Psychological Sequelae")
        sequelae_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #6d28d9; background: transparent; border: none;")
        sequelae_layout.addWidget(sequelae_lbl)

        sequelae_text = QTextEdit()
        sequelae_text.setPlaceholderText("Document PTSD, dissociation, dysregulation, behavioural patterns...")
        sequelae_text.setStyleSheet(input_style + " background: white;")
        sequelae_text.setMinimumHeight(58)
        sequelae_layout.addWidget(sequelae_text)
        sequelae_layout.addWidget(DragResizeBar(sequelae_text))
        text_widgets.append(sequelae_text)
        subsection_data.append(("Psychological Sequelae", sequelae_text))
        setattr(self, "popup_h8_psychological_sequelae", sequelae_text)

        if CollapsibleSection:
            sequelae_section = CollapsibleSection("Select Applicable Categories", start_collapsed=True)
            sequelae_section.set_content_height(200)
            sequelae_section._min_height = 80
            sequelae_section._max_height = 320

            sequelae_section.set_header_style("""
                QFrame {
                    background: rgba(124, 58, 237, 0.15);
                    border: 1px solid rgba(124, 58, 237, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            sequelae_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #6d28d9;
                    background: transparent;
                    border: none;
                }
            """)

            sequelae_content = QWidget()
            sequelae_content.setStyleSheet("""
                QWidget {
                    background: rgba(245, 243, 255, 0.95);
                    border: 1px solid rgba(124, 58, 237, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            sequelae_cb_layout = QVBoxLayout(sequelae_content)
            sequelae_cb_layout.setContentsMargins(10, 6, 10, 6)
            sequelae_cb_layout.setSpacing(3)

            H8_PSYCHOLOGICAL_SEQUELAE = [
                ("Diagnoses/Symptoms", "#7c3aed", [
                    ("ptsd_cptsd", "PTSD / Complex PTSD"),
                    ("dissociation", "Dissociation"),
                    ("hypervigilance", "Hypervigilance"),
                    ("emotional_dysregulation", "Emotional dysregulation"),
                    ("nightmares_flashbacks", "Nightmares / flashbacks"),
                    ("persistent_anger", "Persistent anger or hostility"),
                ]),
                ("Behavioural Patterns", "#8b5cf6", [
                    ("triggered_aggression", "Aggression when triggered"),
                    ("poor_impulse_control", "Poor impulse control"),
                    ("substance_coping", "Substance use as coping"),
                    ("interpersonal_mistrust", "Interpersonal mistrust"),
                    ("threat_reactivity", "Reactivity to perceived threat"),
                ]),
            ]

            h8_sequelae_checkboxes = []
            for cat_name, cat_color, items in H8_PSYCHOLOGICAL_SEQUELAE:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                sequelae_cb_layout.addWidget(cat_lbl)
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("trauma_key", item_key)
                    cb.setProperty("trauma_type", "sequelae")
                    cb.stateChanged.connect(self._update_h8_trauma_narrative)
                    sequelae_cb_layout.addWidget(cb)
                    h8_sequelae_checkboxes.append(cb)

            sequelae_section.set_content(sequelae_content)
            sequelae_layout.addWidget(sequelae_section)
            setattr(self, "popup_h8_sequelae_checkboxes", h8_sequelae_checkboxes)

        if CollapsibleSection:
            subsections_content_layout.addWidget(sequelae_container)
        else:
            layout.addWidget(sequelae_container)

        # === TRAUMA NARRATIVES CONTAINER ===
        narratives_container = QFrame()
        narratives_container.setStyleSheet("""
            QFrame {
                background: rgba(236, 253, 245, 0.5);
                border: 2px solid rgba(5, 150, 105, 0.4);
                border-radius: 10px;
            }
        """)
        narratives_layout = QVBoxLayout(narratives_container)
        narratives_layout.setContentsMargins(12, 10, 12, 10)
        narratives_layout.setSpacing(8)

        narratives_lbl = QLabel("Trauma Narratives")
        narratives_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #047857; background: transparent; border: none;")
        narratives_layout.addWidget(narratives_lbl)

        narratives_text = QTextEdit()
        narratives_text.setPlaceholderText("Document grievance themes, victim identity, persecutory beliefs...")
        narratives_text.setStyleSheet(input_style + " background: white;")
        narratives_text.setMinimumHeight(58)
        narratives_layout.addWidget(narratives_text)
        narratives_layout.addWidget(DragResizeBar(narratives_text))
        text_widgets.append(narratives_text)
        subsection_data.append(("Trauma Narratives", narratives_text))
        setattr(self, "popup_h8_trauma_narratives", narratives_text)

        if CollapsibleSection:
            narratives_section = CollapsibleSection("Select Applicable Themes", start_collapsed=True)
            narratives_section.set_content_height(120)
            narratives_section._min_height = 60
            narratives_section._max_height = 200

            narratives_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.15);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            narratives_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #047857;
                    background: transparent;
                    border: none;
                }
            """)

            narratives_content = QWidget()
            narratives_content.setStyleSheet("""
                QWidget {
                    background: rgba(236, 253, 245, 0.95);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            narratives_cb_layout = QVBoxLayout(narratives_content)
            narratives_cb_layout.setContentsMargins(10, 6, 10, 6)
            narratives_cb_layout.setSpacing(3)

            H8_TRAUMA_NARRATIVES = [
                ("Grievance/Victim Identity", "#059669", [
                    ("everyone_hurts", "\"Everyone has always hurt me\""),
                    ("fight_survive", "\"I had to fight to survive\""),
                    ("cant_trust", "\"You can't trust anyone\""),
                    ("system_abuse", "\"I was treated unfairly / abused by systems\""),
                    ("grievance_identity", "Strong grievance or victim identity"),
                ]),
            ]

            h8_narratives_checkboxes = []
            for cat_name, cat_color, items in H8_TRAUMA_NARRATIVES:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                narratives_cb_layout.addWidget(cat_lbl)
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("trauma_key", item_key)
                    cb.setProperty("trauma_type", "narratives")
                    cb.stateChanged.connect(self._update_h8_trauma_narrative)
                    narratives_cb_layout.addWidget(cb)
                    h8_narratives_checkboxes.append(cb)

            narratives_section.set_content(narratives_content)
            narratives_layout.addWidget(narratives_section)
            setattr(self, "popup_h8_narratives_checkboxes", h8_narratives_checkboxes)

        if CollapsibleSection:
            subsections_content_layout.addWidget(narratives_container)
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)
        else:
            layout.addWidget(narratives_container)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            for label_text, text_widget in subsection_data:
                evidence = text_widget.toPlainText().strip()
                if evidence:
                    parts.append(f"\n{label_text}\n{evidence}")

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high,
        ] + text_widgets)
        self._add_send_button(layout, key, generate)

    def _update_h8_trauma_narrative(self):
        """Update H8 text fields with gender-sensitive narrative based on checked trauma items."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return items[0] + " and " + items[1]
            else:
                return ", ".join(items[:-1]) + ", and " + items[-1]

        # Get text fields
        childhood_text = getattr(self, "popup_h8_childhood_trauma", None)
        adult_text = getattr(self, "popup_h8_adult_trauma", None)
        loss_text = getattr(self, "popup_h8_loss_catastrophe", None)
        sequelae_text = getattr(self, "popup_h8_psychological_sequelae", None)
        narratives_text = getattr(self, "popup_h8_trauma_narratives", None)

        # Get checkboxes
        childhood_cbs = getattr(self, "popup_h8_childhood_checkboxes", [])
        adult_cbs = getattr(self, "popup_h8_adult_checkboxes", [])
        loss_cbs = getattr(self, "popup_h8_loss_checkboxes", [])
        sequelae_cbs = getattr(self, "popup_h8_sequelae_checkboxes", [])
        narratives_cbs = getattr(self, "popup_h8_narratives_checkboxes", [])

        # === CHILDHOOD TRAUMA - Subsection-aware narrative ===
        abuse_keys = ["physical_abuse", "sexual_abuse_child", "emotional_abuse", "witnessed_dv"]
        neglect_keys = ["emotional_neglect", "physical_neglect", "inconsistent_care", "institutional_care", "parental_abandonment"]
        adverse_keys = ["chaotic_household", "parental_substance", "parental_mental_illness", "criminal_caregivers", "placement_breakdowns"]

        abuse_phrases = {"physical_abuse": "physical abuse", "sexual_abuse_child": "sexual abuse", "emotional_abuse": "emotional or psychological abuse", "witnessed_dv": "witnessing domestic violence"}
        neglect_phrases = {"emotional_neglect": "emotional neglect", "physical_neglect": "physical neglect", "inconsistent_care": "inconsistent or absent caregiving", "institutional_care": "instability in institutional or foster care", "parental_abandonment": "parental abandonment"}
        adverse_phrases = {"chaotic_household": "a chaotic household environment", "parental_substance": "parental substance misuse", "parental_mental_illness": "parental mental illness", "criminal_caregivers": "criminal or violent caregivers", "placement_breakdowns": "repeated placement breakdowns"}

        abuse_items, neglect_items, adverse_items = [], [], []
        for cb in childhood_cbs:
            if cb.isChecked():
                key = cb.property("trauma_key")
                if key in abuse_keys and key in abuse_phrases:
                    abuse_items.append(abuse_phrases[key])
                elif key in neglect_keys and key in neglect_phrases:
                    neglect_items.append(neglect_phrases[key])
                elif key in adverse_keys and key in adverse_phrases:
                    adverse_items.append(adverse_phrases[key])

        if childhood_text:
            sentences = []
            if abuse_items:
                sentences.append(f"As a child, {subj} experienced abuse, specifically {join_items(abuse_items)}.")
            if neglect_items:
                if sentences:
                    sentences.append(f"There was also childhood {join_items(neglect_items)}.")
                else:
                    sentences.append(f"As a child, {subj} experienced {join_items(neglect_items)}.")
            if adverse_items:
                if sentences:
                    sentences.append(f"Other adverse events in {poss} upbringing included {join_items(adverse_items)}.")
                else:
                    sentences.append(f"{Poss} upbringing was characterised by {join_items(adverse_items)}.")
            if sentences:
                childhood_text.setPlainText(" ".join(sentences))
            else:
                childhood_text.clear()

        # === ADULT TRAUMA - Subsection-aware narrative ===
        victimisation_keys = ["adult_assault", "sexual_assault_adult", "robbery_violence", "stalking_coercion"]
        institutional_keys = ["prison_violence", "segregation_isolation", "hospital_victimisation", "bullying_harassment"]
        exposure_keys = ["occupational_violence", "witnessed_death", "threats_to_life"]

        victimisation_phrases = {"adult_assault": "assault", "sexual_assault_adult": "sexual assault", "robbery_violence": "robbery with violence", "stalking_coercion": "stalking or coercive control"}
        institutional_phrases = {"prison_violence": "violence in prison", "segregation_isolation": "prolonged segregation or isolation", "hospital_victimisation": "victimisation in hospital or care settings", "bullying_harassment": "bullying, harassment, or exploitation"}
        exposure_phrases = {"occupational_violence": "occupational exposure to violence", "witnessed_death": "witnessing serious injury or death", "threats_to_life": "threats to life"}

        victim_items, inst_items, exp_items = [], [], []
        for cb in adult_cbs:
            if cb.isChecked():
                key = cb.property("trauma_key")
                if key in victimisation_keys and key in victimisation_phrases:
                    victim_items.append(victimisation_phrases[key])
                elif key in institutional_keys and key in institutional_phrases:
                    inst_items.append(institutional_phrases[key])
                elif key in exposure_keys and key in exposure_phrases:
                    exp_items.append(exposure_phrases[key])

        if adult_text:
            sentences = []
            if victim_items:
                sentences.append(f"In adulthood, {subj} has been a victim of {join_items(victim_items)}.")
            if inst_items:
                if sentences:
                    sentences.append(f"{Subj} has also experienced institutional trauma including {join_items(inst_items)}.")
                else:
                    sentences.append(f"{Subj} has experienced institutional trauma including {join_items(inst_items)}.")
            if exp_items:
                if sentences:
                    sentences.append(f"Additionally, {subj} has experienced {join_items(exp_items)}.")
                else:
                    sentences.append(f"{Subj} has experienced {join_items(exp_items)}.")
            if sentences:
                adult_text.setPlainText(" ".join(sentences))
            else:
                adult_text.clear()

        # === LOSS & CATASTROPHE - Subsection-aware narrative ===
        bereavement_keys = ["violent_death", "multiple_bereavements"]
        catastrophe_keys = ["war_displacement", "forced_migration", "serious_accidents", "disasters"]

        bereavement_phrases = {"violent_death": "sudden or violent death of close others", "multiple_bereavements": "multiple bereavements"}
        catastrophe_phrases = {"war_displacement": "war, displacement, or torture", "forced_migration": "forced migration or asylum-related trauma", "serious_accidents": "serious accidents", "disasters": "natural or man-made disasters"}

        bereave_items, catastrophe_items = [], []
        for cb in loss_cbs:
            if cb.isChecked():
                key = cb.property("trauma_key")
                if key in bereavement_keys and key in bereavement_phrases:
                    bereave_items.append(bereavement_phrases[key])
                elif key in catastrophe_keys and key in catastrophe_phrases:
                    catastrophe_items.append(catastrophe_phrases[key])

        if loss_text:
            sentences = []
            if bereave_items:
                sentences.append(f"{Subj} has experienced significant loss including {join_items(bereave_items)}.")
            if catastrophe_items:
                if sentences:
                    sentences.append(f"{Subj} has also been exposed to catastrophic events such as {join_items(catastrophe_items)}.")
                else:
                    sentences.append(f"{Subj} has been exposed to catastrophic events including {join_items(catastrophe_items)}.")
            if sentences:
                loss_text.setPlainText(" ".join(sentences))
            else:
                loss_text.clear()

        # === PSYCHOLOGICAL SEQUELAE - Subsection-aware narrative ===
        diagnosis_keys = ["ptsd_cptsd", "dissociation"]
        arousal_keys = ["hypervigilance", "emotional_dysregulation", "nightmares_flashbacks"]
        behaviour_keys = ["persistent_anger", "triggered_aggression", "poor_impulse_control"]
        coping_keys = ["substance_coping", "interpersonal_mistrust", "threat_reactivity"]

        diagnosis_phrases = {"ptsd_cptsd": "PTSD or Complex PTSD", "dissociation": "dissociative symptoms"}
        arousal_phrases = {"hypervigilance": "hypervigilance", "emotional_dysregulation": "emotional dysregulation", "nightmares_flashbacks": "nightmares or flashbacks"}
        behaviour_phrases = {"persistent_anger": "persistent anger or hostility", "triggered_aggression": "aggression when triggered", "poor_impulse_control": "poor impulse control"}
        coping_phrases = {"substance_coping": "substance use as a coping mechanism", "interpersonal_mistrust": "interpersonal mistrust", "threat_reactivity": "reactivity to perceived threats"}

        diag_items, arousal_items, behav_items, cope_items = [], [], [], []
        for cb in sequelae_cbs:
            if cb.isChecked():
                key = cb.property("trauma_key")
                if key in diagnosis_keys and key in diagnosis_phrases:
                    diag_items.append(diagnosis_phrases[key])
                elif key in arousal_keys and key in arousal_phrases:
                    arousal_items.append(arousal_phrases[key])
                elif key in behaviour_keys and key in behaviour_phrases:
                    behav_items.append(behaviour_phrases[key])
                elif key in coping_keys and key in coping_phrases:
                    cope_items.append(coping_phrases[key])

        if sequelae_text:
            sentences = []
            if diag_items:
                sentences.append(f"As a consequence of trauma, {subj} has been diagnosed with {join_items(diag_items)}.")
            if arousal_items:
                if sentences:
                    sentences.append(f"{Subj} demonstrates trauma-related arousal symptoms including {join_items(arousal_items)}.")
                else:
                    sentences.append(f"{Subj} demonstrates {join_items(arousal_items)}.")
            if behav_items:
                if sentences:
                    sentences.append(f"Behaviourally, {subj} displays {join_items(behav_items)}.")
                else:
                    sentences.append(f"{Subj} displays {join_items(behav_items)}.")
            if cope_items:
                if sentences:
                    sentences.append(f"{Poss} maladaptive coping includes {join_items(cope_items)}.")
                else:
                    sentences.append(f"{Subj} demonstrates {join_items(cope_items)}.")
            if sentences:
                sequelae_text.setPlainText(" ".join(sentences))
            else:
                sequelae_text.clear()

        # === TRAUMA NARRATIVES - Subsection-aware narrative ===
        narrative_phrases = {
            "everyone_hurts": f"believes that everyone has always hurt {obj}",
            "fight_survive": "describes having to fight to survive",
            "cant_trust": "maintains that no one can be trusted",
            "system_abuse": "feels treated unfairly or abused by systems",
            "grievance_identity": "has a strong grievance or victim identity",
        }

        narr_items = []
        for cb in narratives_cbs:
            if cb.isChecked():
                key = cb.property("trauma_key")
                if key in narrative_phrases:
                    narr_items.append(narrative_phrases[key])

        if narratives_text:
            if narr_items:
                narratives_text.setPlainText(f"{Poss} trauma narrative is characterised by the fact that {subj} {join_items(narr_items)}.")
            else:
                narratives_text.clear()

    def _build_hcr_h9_popup(self, key: str, code: str, description: str, subsections: list):
        """Build H9 (Violent Attitudes) popup with operationalized attitude categories.

        Based on HCR-20V3 manual guidance for H9 - enduring cognitive styles that
        support, excuse, justify, minimise, or normalise violence.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"
        section_label_style = "font-size: 16px; font-weight: 600; color: #059669; margin-top: 8px;"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        # === VIOLENT ATTITUDES CONTAINER (Text + Categories) ===
        text_widgets = []
        subsection_data = []

        # --- Subsections CollapsibleSection wrapper ---
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # Violent Attitudes container
        violent_container = QFrame()
        violent_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        violent_container_layout = QVBoxLayout(violent_container)
        violent_container_layout.setContentsMargins(12, 10, 12, 10)
        violent_container_layout.setSpacing(8)

        # Violent Attitudes text input
        violent_lbl = QLabel("Violent Attitudes")
        violent_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        violent_container_layout.addWidget(violent_lbl)

        violent_text = QTextEdit()
        violent_text.setPlaceholderText("Describe attitudes that support, justify, minimise, or normalise violence...")
        violent_text.setStyleSheet(input_style + " background: white;")
        violent_text.setMinimumHeight(58)
        violent_container_layout.addWidget(violent_text)
        violent_container_layout.addWidget(DragResizeBar(violent_text))
        text_widgets.append(violent_text)
        subsection_data.append(("Violent Attitudes", violent_text))
        setattr(self, "popup_h9_violent_attitudes", violent_text)

        # Violent Attitudes Categories (collapsible)
        if CollapsibleSection:
            violent_section = CollapsibleSection("Select Applicable Categories", start_collapsed=False)
            violent_section.set_content_height(260)
            violent_section._min_height = 100
            violent_section._max_height = 400
            violent_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            violent_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            violent_content = QWidget()
            violent_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            violent_cb_layout = QVBoxLayout(violent_content)
            violent_cb_layout.setContentsMargins(10, 6, 10, 6)
            violent_cb_layout.setSpacing(3)

            H9_VIOLENT_ATTITUDES = [
                ("Justification/Endorsement", "#dc2626", [
                    ("justifies_violence", "Views violence as acceptable or justified"),
                    ("victim_deserved", "Believes victim deserved or provoked violence"),
                    ("no_choice", "Claims to have had no choice but to be violent"),
                    ("violence_solves", "Sees violence as effective problem-solving"),
                ]),
                ("Minimisation/Denial", "#ea580c", [
                    ("minimises_harm", "Minimises harm caused to victims"),
                    ("downplays_severity", "Downplays seriousness of violent incidents"),
                    ("denies_intent", "Denies violent intent despite evidence"),
                ]),
                ("Externalisation/Blame", "#f59e0b", [
                    ("blames_victim", "Blames victim for provoking violence"),
                    ("blames_others", "Blames staff, system, or circumstances"),
                    ("external_locus", "Attributes violence to external factors"),
                ]),
                ("Grievance/Entitlement", "#84cc16", [
                    ("feels_persecuted", "Expresses persistent sense of persecution"),
                    ("entitled_respect", "Believes entitled to respect through force"),
                    ("holds_grudges", "Maintains grudges against specific individuals"),
                ]),
                ("Lack of Remorse/Empathy", "#7c3aed", [
                    ("no_remorse", "Shows no remorse for violent behaviour"),
                    ("indifferent_harm", "Appears indifferent to victim suffering"),
                    ("dismisses_impact", "Dismisses impact of violence on others"),
                ]),
            ]

            h9_violent_checkboxes = []

            for cat_name, cat_color, items in H9_VIOLENT_ATTITUDES:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                violent_cb_layout.addWidget(cat_lbl)

                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("attitude_key", item_key)
                    cb.setProperty("attitude_type", "violent")
                    cb.stateChanged.connect(self._update_h9_attitudes_narrative)
                    violent_cb_layout.addWidget(cb)
                    h9_violent_checkboxes.append(cb)

            violent_section.set_content(violent_content)
            violent_container_layout.addWidget(violent_section)
            setattr(self, "popup_h9_violent_checkboxes", h9_violent_checkboxes)

        subsections_target_layout.addWidget(violent_container)

        # === ANTISOCIAL ATTITUDES CONTAINER (Text + Categories) ===
        antisocial_container = QFrame()
        antisocial_container.setStyleSheet("""
            QFrame {
                background: rgba(255, 237, 213, 0.5);
                border: 2px solid rgba(249, 115, 22, 0.4);
                border-radius: 10px;
            }
        """)
        antisocial_container_layout = QVBoxLayout(antisocial_container)
        antisocial_container_layout.setContentsMargins(12, 10, 12, 10)
        antisocial_container_layout.setSpacing(8)

        # Antisocial Attitudes text input
        antisocial_lbl = QLabel("Antisocial Attitudes")
        antisocial_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #c2410c; background: transparent; border: none;")
        antisocial_container_layout.addWidget(antisocial_lbl)

        antisocial_text = QTextEdit()
        antisocial_text.setPlaceholderText("Describe antisocial attitudes, criminal thinking, hostility to authority...")
        antisocial_text.setStyleSheet(input_style + " background: white;")
        antisocial_text.setMinimumHeight(58)
        antisocial_container_layout.addWidget(antisocial_text)
        antisocial_container_layout.addWidget(DragResizeBar(antisocial_text))
        text_widgets.append(antisocial_text)
        subsection_data.append(("Antisocial Attitudes", antisocial_text))
        setattr(self, "popup_h9_antisocial_attitudes", antisocial_text)

        # Antisocial Attitudes Categories (collapsible)
        if CollapsibleSection:
            antisocial_section = CollapsibleSection("Select Applicable Categories", start_collapsed=False)
            antisocial_section.set_content_height(230)
            antisocial_section._min_height = 100
            antisocial_section._max_height = 350
            antisocial_section.set_header_style("""
                QFrame {
                    background: rgba(249, 115, 22, 0.15);
                    border: 1px solid rgba(249, 115, 22, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            antisocial_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #c2410c;
                    background: transparent;
                    border: none;
                }
            """)

            antisocial_content = QWidget()
            antisocial_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 247, 237, 0.95);
                    border: 1px solid rgba(249, 115, 22, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            antisocial_cb_layout = QVBoxLayout(antisocial_content)
            antisocial_cb_layout.setContentsMargins(10, 6, 10, 6)
            antisocial_cb_layout.setSpacing(3)

            H9_ANTISOCIAL_ATTITUDES = [
                ("Criminal Thinking", "#c2410c", [
                    ("rules_dont_apply", "Believes rules or laws do not apply to them"),
                    ("entitled_take", "Feels entitled to take what they want"),
                    ("exploits_others", "Views exploitation of others as acceptable"),
                ]),
                ("Authority Hostility", "#9a3412", [
                    ("hostile_authority", "Expresses hostility toward authority figures"),
                    ("system_corrupt", "Views police, courts, or system as corrupt"),
                    ("staff_deserve", "Believes staff deserve mistreatment"),
                ]),
                ("Callousness", "#78350f", [
                    ("lacks_empathy", "Demonstrates lack of empathy for others"),
                    ("indifferent_consequences", "Indifferent to consequences for others"),
                    ("uses_others", "Views others as means to an end"),
                ]),
                ("Treatment Resistance", "#713f12", [
                    ("rejects_help", "Rejects need for help or intervention"),
                    ("superficial_compliance", "Shows only superficial compliance"),
                    ("unchanged_beliefs", "Beliefs remain unchanged despite treatment"),
                ]),
            ]

            h9_antisocial_checkboxes = []

            for cat_name, cat_color, items in H9_ANTISOCIAL_ATTITUDES:
                cat_lbl = QLabel(cat_name)
                cat_lbl.setStyleSheet(f"font-size: 12px; font-weight: 700; color: {cat_color}; margin-top: 3px; background: transparent;")
                antisocial_cb_layout.addWidget(cat_lbl)

                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("attitude_key", item_key)
                    cb.setProperty("attitude_type", "antisocial")
                    cb.stateChanged.connect(self._update_h9_attitudes_narrative)
                    antisocial_cb_layout.addWidget(cb)
                    h9_antisocial_checkboxes.append(cb)

            antisocial_section.set_content(antisocial_content)
            antisocial_container_layout.addWidget(antisocial_section)
            setattr(self, "popup_h9_antisocial_checkboxes", h9_antisocial_checkboxes)

        subsections_target_layout.addWidget(antisocial_container)

        # --- Close Subsections CollapsibleSection ---
        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            for label_text, text_widget in subsection_data:
                evidence = text_widget.toPlainText().strip()
                if evidence:
                    parts.append(f"\n{label_text}\n{evidence}")

            return "\n".join(parts)

        self._connect_preview_updates(key, [
            presence_no, presence_partial, presence_yes, presence_omit,
            relevance_low, relevance_mod, relevance_high,
        ] + text_widgets)
        self._add_send_button(layout, key, generate)

    def _update_h9_attitudes_narrative(self):
        """Update H9 text fields with subsection-aware, gender-sensitive narrative."""
        # Get gender for pronouns
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # Get text fields
        violent_text = getattr(self, "popup_h9_violent_attitudes", None)
        antisocial_text = getattr(self, "popup_h9_antisocial_attitudes", None)

        # Get checkboxes
        violent_cbs = getattr(self, "popup_h9_violent_checkboxes", [])
        antisocial_cbs = getattr(self, "popup_h9_antisocial_checkboxes", [])

        # ============ VIOLENT ATTITUDES - Subsection aware ============
        # Subsection mappings with item labels
        justification_items = {
            "justifies_violence": "viewing violence as acceptable",
            "victim_deserved": "believing victims provoked or deserved it",
            "no_choice": "claiming to have had no choice",
            "violence_solves": "seeing violence as effective problem-solving",
        }
        minimisation_items = {
            "minimises_harm": "minimising harm caused",
            "downplays_severity": "downplaying the seriousness of incidents",
            "denies_intent": "denying intent despite evidence",
        }
        externalisation_items = {
            "blames_victim": "blaming victims",
            "blames_others": "blaming staff or circumstances",
            "external_locus": "attributing violence to factors beyond {poss} control",
        }
        grievance_items = {
            "feels_persecuted": "a persistent sense of persecution",
            "entitled_respect": "believing {subj} is entitled to respect through force",
            "holds_grudges": "maintaining grudges against specific individuals",
        }
        remorse_items = {
            "no_remorse": "no remorse for {poss} violent behaviour",
            "indifferent_harm": "indifference to victim suffering",
            "dismisses_impact": "dismissing the impact on others",
        }

        # Collect checked items by subsection
        justification = []
        minimisation = []
        externalisation = []
        grievance = []
        remorse = []

        for cb in violent_cbs:
            if cb.isChecked():
                key = cb.property("attitude_key")
                if key in justification_items:
                    justification.append(justification_items[key])
                elif key in minimisation_items:
                    minimisation.append(minimisation_items[key])
                elif key in externalisation_items:
                    externalisation.append(externalisation_items[key].format(poss=poss, subj=subj))
                elif key in grievance_items:
                    grievance.append(grievance_items[key].format(poss=poss, subj=subj))
                elif key in remorse_items:
                    remorse.append(remorse_items[key].format(poss=poss, subj=subj))

        if violent_text:
            sentences = []

            if justification:
                sentences.append(f"{Subj} demonstrates justification and endorsement of violence, {join_items(justification)}.")

            if minimisation:
                if sentences:
                    sentences.append(f"{Subj} also engages in minimisation by {join_items(minimisation)}.")
                else:
                    sentences.append(f"{Subj} engages in minimisation of {poss} violence, {join_items(minimisation)}.")

            if externalisation:
                if sentences:
                    sentences.append(f"There is also externalisation of blame, with {obj} {join_items(externalisation)}.")
                else:
                    sentences.append(f"{Subj} externalises blame for {poss} violence, {join_items(externalisation)}.")

            if grievance:
                if sentences:
                    sentences.append(f"Grievance thinking is evident, including {join_items(grievance)}.")
                else:
                    sentences.append(f"{Subj} demonstrates grievance thinking, including {join_items(grievance)}.")

            if remorse:
                if sentences:
                    sentences.append(f"Finally, {subj} shows {join_items(remorse)}.")
                else:
                    sentences.append(f"{Subj} shows {join_items(remorse)}.")

            if sentences:
                violent_text.setPlainText(" ".join(sentences))
            else:
                violent_text.clear()

        # ============ ANTISOCIAL ATTITUDES - Subsection aware ============
        criminal_items = {
            "rules_dont_apply": "believing rules and laws do not apply to {obj}",
            "entitled_take": "feeling entitled to take what {subj} wants",
            "exploits_others": "viewing exploitation of others as acceptable",
        }
        authority_items = {
            "hostile_authority": "hostility toward authority figures",
            "system_corrupt": "viewing the system as corrupt or biased",
            "staff_deserve": "believing staff deserve mistreatment",
        }
        callousness_items = {
            "lacks_empathy": "a lack of empathy",
            "indifferent_consequences": "indifference to consequences for others",
            "uses_others": "viewing others as a means to an end",
        }
        resistance_items = {
            "rejects_help": "rejecting the need for help",
            "superficial_compliance": "superficial compliance only",
            "unchanged_beliefs": "beliefs remaining unchanged despite treatment",
        }

        # Collect checked items by subsection
        criminal = []
        authority = []
        callousness = []
        resistance = []

        for cb in antisocial_cbs:
            if cb.isChecked():
                key = cb.property("attitude_key")
                if key in criminal_items:
                    criminal.append(criminal_items[key].format(obj=obj, subj=subj))
                elif key in authority_items:
                    authority.append(authority_items[key])
                elif key in callousness_items:
                    callousness.append(callousness_items[key])
                elif key in resistance_items:
                    resistance.append(resistance_items[key])

        if antisocial_text:
            sentences = []

            if criminal:
                sentences.append(f"{Subj} demonstrates criminal thinking patterns, {join_items(criminal)}.")

            if authority:
                if sentences:
                    sentences.append(f"{Subj} also shows {join_items(authority)}.")
                else:
                    sentences.append(f"{Subj} demonstrates {join_items(authority)}.")

            if callousness:
                if sentences:
                    sentences.append(f"There is also evidence of callousness, including {join_items(callousness)}.")
                else:
                    sentences.append(f"{Subj} demonstrates callousness, including {join_items(callousness)}.")

            if resistance:
                if sentences:
                    sentences.append(f"With respect to treatment, {subj} shows {join_items(resistance)}.")
                else:
                    sentences.append(f"{Subj} shows treatment resistance, including {join_items(resistance)}.")

            if sentences:
                antisocial_text.setPlainText(" ".join(sentences))
            else:
                antisocial_text.clear()

    def _build_hcr_h10_popup(self, key: str, code: str, description: str, subsections: list):
        """Build H10 (Treatment/Supervision Response) popup with treatment response categories.

        Based on HCR-20V3 manual guidance for H10 - problems engaging with,
        complying with, and responding to treatment and supervision.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        # Store text widgets and subsection data
        text_widgets = []
        subsection_data = []

        # --- Subsections CollapsibleSection wrapper ---
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)

            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. MEDICATION NON-ADHERENCE CONTAINER ===
        med_container = QFrame()
        med_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        med_container_layout = QVBoxLayout(med_container)
        med_container_layout.setContentsMargins(12, 10, 12, 10)
        med_container_layout.setSpacing(8)

        med_lbl = QLabel("Medication Non-Adherence")
        med_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        med_container_layout.addWidget(med_lbl)

        med_text = QTextEdit()
        med_text.setPlaceholderText("Describe patterns of medication non-compliance, refusal, or poor adherence...")
        med_text.setStyleSheet(input_style + " background: white;")
        med_text.setMinimumHeight(58)
        med_container_layout.addWidget(med_text)
        med_container_layout.addWidget(DragResizeBar(med_text))
        text_widgets.append(med_text)
        subsection_data.append(("Medication Non-Adherence", med_text))
        setattr(self, "popup_h10_medication", med_text)

        if CollapsibleSection:
            med_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            med_section.set_content_height(200)
            med_section._min_height = 80
            med_section._max_height = 300
            med_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            med_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            med_content = QWidget()
            med_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            med_cb_layout = QVBoxLayout(med_content)
            med_cb_layout.setContentsMargins(10, 6, 10, 6)
            med_cb_layout.setSpacing(3)

            H10_MEDICATION_ITEMS = [
                ("med_noncompliant", "Non-compliant with medication"),
                ("med_poor_adherence", "Poor adherence to prescribed treatment"),
                ("med_frequent_refusal", "Frequently refuses medication"),
                ("med_stopped_without", "Stopped medication without medical advice"),
                ("med_intermittent", "Intermittent compliance"),
                ("med_refused_depot", "Refused depot injection"),
                ("med_self_discontinued", "Self-discontinued medication"),
                ("med_repeated_stopping", "Repeated stopping/starting pattern"),
            ]

            h10_med_checkboxes = []
            for item_key, item_label in H10_MEDICATION_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("h10_key", item_key)
                cb.setProperty("h10_category", "medication")
                cb.stateChanged.connect(self._update_h10_treatment_narrative)
                med_cb_layout.addWidget(cb)
                h10_med_checkboxes.append(cb)

            med_section.set_content(med_content)
            med_container_layout.addWidget(med_section)
            setattr(self, "popup_h10_med_checkboxes", h10_med_checkboxes)

        subsections_target_layout.addWidget(med_container)

        # === 2. DISENGAGEMENT FROM SERVICES CONTAINER ===
        disengage_container = QFrame()
        disengage_container.setStyleSheet("""
            QFrame {
                background: rgba(255, 237, 213, 0.5);
                border: 2px solid rgba(249, 115, 22, 0.4);
                border-radius: 10px;
            }
        """)
        disengage_container_layout = QVBoxLayout(disengage_container)
        disengage_container_layout.setContentsMargins(12, 10, 12, 10)
        disengage_container_layout.setSpacing(8)

        disengage_lbl = QLabel("Disengagement From Services")
        disengage_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #c2410c; background: transparent; border: none;")
        disengage_container_layout.addWidget(disengage_lbl)

        disengage_text = QTextEdit()
        disengage_text.setPlaceholderText("Describe patterns of disengagement, DNAs, poor engagement with services...")
        disengage_text.setStyleSheet(input_style + " background: white;")
        disengage_text.setMinimumHeight(58)
        disengage_container_layout.addWidget(disengage_text)
        disengage_container_layout.addWidget(DragResizeBar(disengage_text))
        text_widgets.append(disengage_text)
        subsection_data.append(("Disengagement From Services", disengage_text))
        setattr(self, "popup_h10_disengagement", disengage_text)

        if CollapsibleSection:
            disengage_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            disengage_section.set_content_height(200)
            disengage_section._min_height = 80
            disengage_section._max_height = 300
            disengage_section.set_header_style("""
                QFrame {
                    background: rgba(249, 115, 22, 0.15);
                    border: 1px solid rgba(249, 115, 22, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            disengage_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #c2410c;
                    background: transparent;
                    border: none;
                }
            """)

            disengage_content = QWidget()
            disengage_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 247, 237, 0.95);
                    border: 1px solid rgba(249, 115, 22, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            disengage_cb_layout = QVBoxLayout(disengage_content)
            disengage_cb_layout.setContentsMargins(10, 6, 10, 6)
            disengage_cb_layout.setSpacing(3)

            H10_DISENGAGEMENT_ITEMS = [
                ("dis_dna", "DNA appointments repeatedly"),
                ("dis_disengaged", "Disengaged from services"),
                ("dis_lost_followup", "Lost to follow-up"),
                ("dis_poor_engagement", "Poor engagement with care team"),
                ("dis_minimal_mdt", "Minimal engagement with MDT"),
                ("dis_refuses_community", "Refuses community follow-up"),
                ("dis_uncontactable", "Uncontactable for prolonged periods"),
            ]

            h10_disengage_checkboxes = []
            for item_key, item_label in H10_DISENGAGEMENT_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("h10_key", item_key)
                cb.setProperty("h10_category", "disengagement")
                cb.stateChanged.connect(self._update_h10_treatment_narrative)
                disengage_cb_layout.addWidget(cb)
                h10_disengage_checkboxes.append(cb)

            disengage_section.set_content(disengage_content)
            disengage_container_layout.addWidget(disengage_section)
            setattr(self, "popup_h10_disengage_checkboxes", h10_disengage_checkboxes)

        subsections_target_layout.addWidget(disengage_container)

        # === 3. RESISTANCE/HOSTILITY TOWARD TREATMENT CONTAINER ===
        hostile_container = QFrame()
        hostile_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 249, 195, 0.5);
                border: 2px solid rgba(202, 138, 4, 0.4);
                border-radius: 10px;
            }
        """)
        hostile_container_layout = QVBoxLayout(hostile_container)
        hostile_container_layout.setContentsMargins(12, 10, 12, 10)
        hostile_container_layout.setSpacing(8)

        hostile_lbl = QLabel("Resistance or Hostility Toward Treatment")
        hostile_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #a16207; background: transparent; border: none;")
        hostile_container_layout.addWidget(hostile_lbl)

        hostile_text = QTextEdit()
        hostile_text.setPlaceholderText("Describe attitudinal resistance, hostility to staff, dismissiveness of treatment...")
        hostile_text.setStyleSheet(input_style + " background: white;")
        hostile_text.setMinimumHeight(58)
        hostile_container_layout.addWidget(hostile_text)
        hostile_container_layout.addWidget(DragResizeBar(hostile_text))
        text_widgets.append(hostile_text)
        subsection_data.append(("Resistance/Hostility Toward Treatment", hostile_text))
        setattr(self, "popup_h10_hostility", hostile_text)

        if CollapsibleSection:
            hostile_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            hostile_section.set_content_height(200)
            hostile_section._min_height = 80
            hostile_section._max_height = 300
            hostile_section.set_header_style("""
                QFrame {
                    background: rgba(202, 138, 4, 0.15);
                    border: 1px solid rgba(202, 138, 4, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            hostile_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #a16207;
                    background: transparent;
                    border: none;
                }
            """)

            hostile_content = QWidget()
            hostile_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(202, 138, 4, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            hostile_cb_layout = QVBoxLayout(hostile_content)
            hostile_cb_layout.setContentsMargins(10, 6, 10, 6)
            hostile_cb_layout.setSpacing(3)

            H10_HOSTILITY_ITEMS = [
                ("hos_refuses_engage", "Refuses to engage with treatment"),
                ("hos_hostile_staff", "Hostile to staff"),
                ("hos_dismissive", "Dismissive of treatment"),
                ("hos_no_insight", "Lacks insight into need for treatment"),
                ("hos_not_necessary", "Does not believe treatment is necessary"),
                ("hos_rejects_psych", "Rejects psychological input"),
                ("hos_uncooperative", "Uncooperative with ward rules"),
                ("hos_oppositional", "Oppositional behaviour toward clinicians"),
            ]

            h10_hostile_checkboxes = []
            for item_key, item_label in H10_HOSTILITY_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("h10_key", item_key)
                cb.setProperty("h10_category", "hostility")
                cb.stateChanged.connect(self._update_h10_treatment_narrative)
                hostile_cb_layout.addWidget(cb)
                h10_hostile_checkboxes.append(cb)

            hostile_section.set_content(hostile_content)
            hostile_container_layout.addWidget(hostile_section)
            setattr(self, "popup_h10_hostile_checkboxes", h10_hostile_checkboxes)

        subsections_target_layout.addWidget(hostile_container)

        # === 4. FAILURE UNDER SUPERVISION CONTAINER ===
        failure_container = QFrame()
        failure_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(22, 163, 74, 0.4);
                border-radius: 10px;
            }
        """)
        failure_container_layout = QVBoxLayout(failure_container)
        failure_container_layout.setContentsMargins(12, 10, 12, 10)
        failure_container_layout.setSpacing(8)

        failure_lbl = QLabel("Failure Under Supervision")
        failure_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        failure_container_layout.addWidget(failure_lbl)

        failure_text = QTextEdit()
        failure_text.setPlaceholderText("Describe breaches, recalls, absconding, failed community placements...")
        failure_text.setStyleSheet(input_style + " background: white;")
        failure_text.setMinimumHeight(58)
        failure_container_layout.addWidget(failure_text)
        failure_container_layout.addWidget(DragResizeBar(failure_text))
        text_widgets.append(failure_text)
        subsection_data.append(("Failure Under Supervision", failure_text))
        setattr(self, "popup_h10_failure", failure_text)

        if CollapsibleSection:
            failure_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            failure_section.set_content_height(220)
            failure_section._min_height = 80
            failure_section._max_height = 320
            failure_section.set_header_style("""
                QFrame {
                    background: rgba(22, 163, 74, 0.15);
                    border: 1px solid rgba(22, 163, 74, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            failure_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #15803d;
                    background: transparent;
                    border: none;
                }
            """)

            failure_content = QWidget()
            failure_content.setStyleSheet("""
                QWidget {
                    background: rgba(240, 253, 244, 0.95);
                    border: 1px solid rgba(22, 163, 74, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            failure_cb_layout = QVBoxLayout(failure_content)
            failure_cb_layout.setContentsMargins(10, 6, 10, 6)
            failure_cb_layout.setSpacing(3)

            H10_FAILURE_ITEMS = [
                ("fail_breach_conditions", "Breach of conditions"),
                ("fail_breach_cto", "Breach of CTO"),
                ("fail_breach_probation", "Breach of probation"),
                ("fail_recall", "Recall to hospital"),
                ("fail_returned_custody", "Returned to custody"),
                ("fail_licence_breach", "Non-compliance with licence conditions"),
                ("fail_community_placement", "Failed community placement"),
                ("fail_absconded", "Absconded / AWOL"),
                ("fail_repeated_recalls", "Repeated recalls or breaches"),
            ]

            h10_failure_checkboxes = []
            for item_key, item_label in H10_FAILURE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("h10_key", item_key)
                cb.setProperty("h10_category", "failure")
                cb.stateChanged.connect(self._update_h10_treatment_narrative)
                failure_cb_layout.addWidget(cb)
                h10_failure_checkboxes.append(cb)

            failure_section.set_content(failure_content)
            failure_container_layout.addWidget(failure_section)
            setattr(self, "popup_h10_failure_checkboxes", h10_failure_checkboxes)

        subsections_target_layout.addWidget(failure_container)

        # === 5. INEFFECTIVE PAST INTERVENTIONS CONTAINER ===
        ineffective_container = QFrame()
        ineffective_container.setStyleSheet("""
            QFrame {
                background: rgba(224, 231, 255, 0.5);
                border: 2px solid rgba(99, 102, 241, 0.4);
                border-radius: 10px;
            }
        """)
        ineffective_container_layout = QVBoxLayout(ineffective_container)
        ineffective_container_layout.setContentsMargins(12, 10, 12, 10)
        ineffective_container_layout.setSpacing(8)

        ineffective_lbl = QLabel("Ineffective Past Interventions")
        ineffective_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #4338ca; background: transparent; border: none;")
        ineffective_container_layout.addWidget(ineffective_lbl)

        ineffective_text = QTextEdit()
        ineffective_text.setPlaceholderText("Describe interventions that did not work, limited treatment response, relapses despite support...")
        ineffective_text.setStyleSheet(input_style + " background: white;")
        ineffective_text.setMinimumHeight(58)
        ineffective_container_layout.addWidget(ineffective_text)
        ineffective_container_layout.addWidget(DragResizeBar(ineffective_text))
        text_widgets.append(ineffective_text)
        subsection_data.append(("Ineffective Past Interventions", ineffective_text))
        setattr(self, "popup_h10_ineffective", ineffective_text)

        if CollapsibleSection:
            ineffective_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            ineffective_section.set_content_height(200)
            ineffective_section._min_height = 80
            ineffective_section._max_height = 300
            ineffective_section.set_header_style("""
                QFrame {
                    background: rgba(99, 102, 241, 0.15);
                    border: 1px solid rgba(99, 102, 241, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            ineffective_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #4338ca;
                    background: transparent;
                    border: none;
                }
            """)

            ineffective_content = QWidget()
            ineffective_content.setStyleSheet("""
                QWidget {
                    background: rgba(238, 242, 255, 0.95);
                    border: 1px solid rgba(99, 102, 241, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            ineffective_cb_layout = QVBoxLayout(ineffective_content)
            ineffective_cb_layout.setContentsMargins(10, 6, 10, 6)
            ineffective_cb_layout.setSpacing(3)

            H10_INEFFECTIVE_ITEMS = [
                ("inef_little_benefit", "Little benefit from treatment"),
                ("inef_limited_response", "Limited response to interventions"),
                ("inef_no_sustained", "No sustained improvement"),
                ("inef_gains_not_maintained", "Treatment gains not maintained"),
                ("inef_relapse_discharge", "Relapse following discharge"),
                ("inef_risk_escalated", "Risk escalated despite treatment"),
                ("inef_repeated_admissions", "Repeated admissions despite support"),
            ]

            h10_ineffective_checkboxes = []
            for item_key, item_label in H10_INEFFECTIVE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("h10_key", item_key)
                cb.setProperty("h10_category", "ineffective")
                cb.stateChanged.connect(self._update_h10_treatment_narrative)
                ineffective_cb_layout.addWidget(cb)
                h10_ineffective_checkboxes.append(cb)

            ineffective_section.set_content(ineffective_content)
            ineffective_container_layout.addWidget(ineffective_section)
            setattr(self, "popup_h10_ineffective_checkboxes", h10_ineffective_checkboxes)

        subsections_target_layout.addWidget(ineffective_container)

        # === 6. ONLY COMPLIES UNDER COMPULSION CONTAINER ===
        compulsion_container = QFrame()
        compulsion_container.setStyleSheet("""
            QFrame {
                background: rgba(243, 232, 255, 0.5);
                border: 2px solid rgba(147, 51, 234, 0.4);
                border-radius: 10px;
            }
        """)
        compulsion_container_layout = QVBoxLayout(compulsion_container)
        compulsion_container_layout.setContentsMargins(12, 10, 12, 10)
        compulsion_container_layout.setSpacing(8)

        compulsion_lbl = QLabel("Only Complies Under Compulsion")
        compulsion_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #7e22ce; background: transparent; border: none;")
        compulsion_container_layout.addWidget(compulsion_lbl)

        compulsion_text = QTextEdit()
        compulsion_text.setPlaceholderText("Describe compliance only when detained, deterioration in community, need for legal framework...")
        compulsion_text.setStyleSheet(input_style + " background: white;")
        compulsion_text.setMinimumHeight(58)
        compulsion_container_layout.addWidget(compulsion_text)
        compulsion_container_layout.addWidget(DragResizeBar(compulsion_text))
        text_widgets.append(compulsion_text)
        subsection_data.append(("Only Complies Under Compulsion", compulsion_text))
        setattr(self, "popup_h10_compulsion", compulsion_text)

        if CollapsibleSection:
            compulsion_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            compulsion_section.set_content_height(180)
            compulsion_section._min_height = 80
            compulsion_section._max_height = 280
            compulsion_section.set_header_style("""
                QFrame {
                    background: rgba(147, 51, 234, 0.15);
                    border: 1px solid rgba(147, 51, 234, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            compulsion_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #7e22ce;
                    background: transparent;
                    border: none;
                }
            """)

            compulsion_content = QWidget()
            compulsion_content.setStyleSheet("""
                QWidget {
                    background: rgba(250, 245, 255, 0.95);
                    border: 1px solid rgba(147, 51, 234, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            compulsion_cb_layout = QVBoxLayout(compulsion_content)
            compulsion_cb_layout.setContentsMargins(10, 6, 10, 6)
            compulsion_cb_layout.setSpacing(3)

            H10_COMPULSION_ITEMS = [
                ("comp_only_under_section", "Only compliant under section"),
                ("comp_engages_detained", "Engages only when detained"),
                ("comp_deteriorates_community", "Deteriorates in community setting"),
                ("comp_legal_framework", "Compliance contingent on legal framework"),
                ("comp_enforced_only", "Responds only to enforced treatment"),
            ]

            h10_compulsion_checkboxes = []
            for item_key, item_label in H10_COMPULSION_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("h10_key", item_key)
                cb.setProperty("h10_category", "compulsion")
                cb.stateChanged.connect(self._update_h10_treatment_narrative)
                compulsion_cb_layout.addWidget(cb)
                h10_compulsion_checkboxes.append(cb)

            compulsion_section.set_content(compulsion_content)
            compulsion_container_layout.addWidget(compulsion_section)
            setattr(self, "popup_h10_compulsion_checkboxes", h10_compulsion_checkboxes)

        subsections_target_layout.addWidget(compulsion_container)

        # --- Close Subsections CollapsibleSection ---
        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")

            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_h10_treatment_narrative(self):
        """Update H10 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. MEDICATION NON-COMPLIANCE ============
        med_items = {
            "med_noncompliant": "non-compliance with medication",
            "med_poor_adherence": "poor adherence to {poss} treatment regimen",
            "med_frequent_refusal": "frequent refusal of medication",
            "med_stopped_without": "stopping medication without medical advice",
            "med_intermittent": "intermittent compliance",
            "med_refused_depot": "refusal of depot injections",
            "med_self_discontinued": "self-discontinuation of medication",
            "med_repeated_stopping": "a pattern of repeatedly stopping and starting",
        }

        med_cbs = getattr(self, "popup_h10_med_checkboxes", [])
        med_text = getattr(self, "popup_h10_medication", None)

        med_checked = []
        for cb in med_cbs:
            if cb.isChecked():
                key = cb.property("h10_key")
                if key in med_items:
                    med_checked.append(med_items[key].format(poss=poss))

        if med_text:
            if med_checked:
                med_text.setPlainText(f"There is a history of medication non-compliance, including {join_items(med_checked)}.")
            else:
                med_text.clear()

        # ============ 2. DISENGAGEMENT FROM SERVICES ============
        dis_items = {
            "dis_dna": "repeatedly failing to attend appointments",
            "dis_disengaged": "disengaging from services",
            "dis_lost_followup": "being lost to follow-up",
            "dis_poor_engagement": "poor engagement with {poss} care team",
            "dis_minimal_mdt": "minimal engagement with the MDT",
            "dis_refuses_community": "refusing community follow-up",
            "dis_uncontactable": "being uncontactable for prolonged periods",
        }

        dis_cbs = getattr(self, "popup_h10_disengage_checkboxes", [])
        dis_text = getattr(self, "popup_h10_disengagement", None)

        dis_checked = []
        for cb in dis_cbs:
            if cb.isChecked():
                key = cb.property("h10_key")
                if key in dis_items:
                    dis_checked.append(dis_items[key].format(poss=poss))

        if dis_text:
            if dis_checked:
                dis_text.setPlainText(f"{Subj} has demonstrated patterns of disengagement from services, {join_items(dis_checked)}.")
            else:
                dis_text.clear()

        # ============ 3. RESISTANCE/HOSTILITY TOWARD TREATMENT ============
        hos_items = {
            "hos_refuses_engage": "refusing to engage with treatment",
            "hos_hostile_staff": "hostility toward clinical staff",
            "hos_dismissive": "being dismissive of treatment",
            "hos_no_insight": "lacking insight into the need for treatment",
            "hos_not_necessary": "not believing treatment is necessary",
            "hos_rejects_psych": "rejecting psychological input",
            "hos_uncooperative": "being uncooperative with ward rules",
            "hos_oppositional": "oppositional behaviour toward clinicians",
        }

        hos_cbs = getattr(self, "popup_h10_hostile_checkboxes", [])
        hos_text = getattr(self, "popup_h10_hostility", None)

        hos_checked = []
        for cb in hos_cbs:
            if cb.isChecked():
                key = cb.property("h10_key")
                if key in hos_items:
                    hos_checked.append(hos_items[key])

        if hos_text:
            if hos_checked:
                hos_text.setPlainText(f"{Subj} has shown resistance and hostility toward treatment, including {join_items(hos_checked)}.")
            else:
                hos_text.clear()

        # ============ 4. FAILURE UNDER SUPERVISION ============
        fail_items = {
            "fail_breach_conditions": "breach of supervision conditions",
            "fail_breach_cto": "breach of CTO",
            "fail_breach_probation": "breach of probation",
            "fail_recall": "recall to hospital",
            "fail_returned_custody": "return to custody",
            "fail_licence_breach": "non-compliance with licence conditions",
            "fail_community_placement": "failed community placements",
            "fail_absconded": "absconding or going AWOL",
            "fail_repeated_recalls": "repeated recalls and breaches",
        }

        fail_cbs = getattr(self, "popup_h10_failure_checkboxes", [])
        fail_text = getattr(self, "popup_h10_failure", None)

        fail_checked = []
        for cb in fail_cbs:
            if cb.isChecked():
                key = cb.property("h10_key")
                if key in fail_items:
                    fail_checked.append(fail_items[key])

        if fail_text:
            if fail_checked:
                fail_text.setPlainText(f"Supervision failures are documented, including {join_items(fail_checked)}.")
            else:
                fail_text.clear()

        # ============ 5. INEFFECTIVE PAST INTERVENTIONS ============
        inef_items = {
            "inef_little_benefit": "little benefit from previous treatment",
            "inef_limited_response": "limited response to interventions",
            "inef_no_sustained": "no sustained improvement",
            "inef_gains_not_maintained": "treatment gains not being maintained",
            "inef_relapse_discharge": "relapse following discharge",
            "inef_risk_escalated": "risk escalation despite treatment",
            "inef_repeated_admissions": "repeated admissions despite community support",
        }

        inef_cbs = getattr(self, "popup_h10_ineffective_checkboxes", [])
        inef_text = getattr(self, "popup_h10_ineffective", None)

        inef_checked = []
        for cb in inef_cbs:
            if cb.isChecked():
                key = cb.property("h10_key")
                if key in inef_items:
                    inef_checked.append(inef_items[key])

        if inef_text:
            if inef_checked:
                inef_text.setPlainText(f"Past interventions have been of limited effectiveness, with {join_items(inef_checked)}.")
            else:
                inef_text.clear()

        # ============ 6. ONLY COMPLIES UNDER COMPULSION ============
        comp_items = {
            "comp_only_under_section": "being compliant only when under section",
            "comp_engages_detained": "engaging only when detained",
            "comp_deteriorates_community": "deteriorating in community settings",
            "comp_legal_framework": "compliance being contingent on a legal framework",
            "comp_enforced_only": "responding only to enforced treatment",
        }

        comp_cbs = getattr(self, "popup_h10_compulsion_checkboxes", [])
        comp_text = getattr(self, "popup_h10_compulsion", None)

        comp_checked = []
        for cb in comp_cbs:
            if cb.isChecked():
                key = cb.property("h10_key")
                if key in comp_items:
                    comp_checked.append(comp_items[key])

        if comp_text:
            if comp_checked:
                comp_text.setPlainText(f"There is evidence that {subj} only complies under compulsion, with {join_items(comp_checked)}.")
            else:
                comp_text.clear()

    def _build_hcr_c1_popup(self, key: str, code: str, description: str, subsections: list):
        """Build C1 (Insight) popup with insight categories.

        Based on HCR-20V3 manual guidance for C1 - awareness of mental disorder,
        risk factors, need for treatment, and link between illness and violence.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. INSIGHT INTO MENTAL DISORDER CONTAINER ===
        disorder_container = QFrame()
        disorder_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        disorder_layout = QVBoxLayout(disorder_container)
        disorder_layout.setContentsMargins(12, 10, 12, 10)
        disorder_layout.setSpacing(8)

        disorder_lbl = QLabel("Insight into Mental Disorder")
        disorder_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        disorder_layout.addWidget(disorder_lbl)

        disorder_text = QTextEdit()
        disorder_text.setPlaceholderText("Describe awareness and acceptance of mental disorder, diagnosis, symptoms...")
        disorder_text.setStyleSheet(input_style + " background: white;")
        disorder_text.setMinimumHeight(58)
        disorder_layout.addWidget(disorder_text)
        disorder_layout.addWidget(DragResizeBar(disorder_text))
        text_widgets.append(disorder_text)
        subsection_data.append(("Insight into Mental Disorder", disorder_text))
        setattr(self, "popup_c1_disorder", disorder_text)

        if CollapsibleSection:
            disorder_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            disorder_section.set_content_height(200)
            disorder_section._min_height = 80
            disorder_section._max_height = 300
            disorder_section.set_header_style("""
                QFrame {
                    background: rgba(59, 130, 246, 0.15);
                    border: 1px solid rgba(59, 130, 246, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            disorder_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #1d4ed8;
                    background: transparent;
                    border: none;
                }
            """)

            disorder_content = QWidget()
            disorder_content.setStyleSheet("""
                QWidget {
                    background: rgba(239, 246, 255, 0.95);
                    border: 1px solid rgba(59, 130, 246, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            disorder_cb_layout = QVBoxLayout(disorder_content)
            disorder_cb_layout.setContentsMargins(10, 6, 10, 6)
            disorder_cb_layout.setSpacing(3)

            C1_DISORDER_ITEMS = [
                ("dis_denies_illness", "Denies illness / nothing wrong with me"),
                ("dis_rejects_diagnosis", "Rejects diagnosis"),
                ("dis_external_attribution", "Attributes symptoms externally (e.g., staff, system)"),
                ("dis_poor_insight", "Poor/limited insight documented"),
                ("dis_no_recognise_relapse", "Does not recognise relapse signs"),
                ("dis_accepts_diagnosis", "Accepts diagnosis (protective)"),
                ("dis_recognises_symptoms", "Recognises symptoms as illness-related"),
            ]

            c1_disorder_checkboxes = []
            for item_key, item_label in C1_DISORDER_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c1_key", item_key)
                cb.setProperty("c1_category", "disorder")
                cb.stateChanged.connect(self._update_c1_insight_narrative)
                disorder_cb_layout.addWidget(cb)
                c1_disorder_checkboxes.append(cb)

            disorder_section.set_content(disorder_content)
            disorder_layout.addWidget(disorder_section)
            setattr(self, "popup_c1_disorder_checkboxes", c1_disorder_checkboxes)

        subsections_target_layout.addWidget(disorder_container)

        # === 2. INSIGHT INTO ILLNESS-RISK LINK CONTAINER ===
        link_container = QFrame()
        link_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        link_layout = QVBoxLayout(link_container)
        link_layout.setContentsMargins(12, 10, 12, 10)
        link_layout.setSpacing(8)

        link_lbl = QLabel("Insight into Illness-Risk Link")
        link_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        link_layout.addWidget(link_lbl)

        link_text = QTextEdit()
        link_text.setPlaceholderText("Describe understanding of link between mental state and past violence/risk...")
        link_text.setStyleSheet(input_style + " background: white;")
        link_text.setMinimumHeight(58)
        link_layout.addWidget(link_text)
        link_layout.addWidget(DragResizeBar(link_text))
        text_widgets.append(link_text)
        subsection_data.append(("Insight into Illness-Risk Link", link_text))
        setattr(self, "popup_c1_link", link_text)

        if CollapsibleSection:
            link_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            link_section.set_content_height(200)
            link_section._min_height = 80
            link_section._max_height = 300
            link_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            link_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            link_content = QWidget()
            link_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            link_cb_layout = QVBoxLayout(link_content)
            link_cb_layout.setContentsMargins(10, 6, 10, 6)
            link_cb_layout.setSpacing(3)

            C1_LINK_ITEMS = [
                ("link_denies_connection", "Denies link between illness and offending"),
                ("link_minimises_violence", "Minimises past violence"),
                ("link_externalises_blame", "Externalises blame (they provoked me)"),
                ("link_lacks_victim_empathy", "Lacks victim empathy"),
                ("link_no_reflection", "Limited/no reflection on index offence"),
                ("link_understands_triggers", "Understands triggers (protective)"),
                ("link_acknowledges_unwell", "Acknowledges was unwell during offence"),
            ]

            c1_link_checkboxes = []
            for item_key, item_label in C1_LINK_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c1_key", item_key)
                cb.setProperty("c1_category", "link")
                cb.stateChanged.connect(self._update_c1_insight_narrative)
                link_cb_layout.addWidget(cb)
                c1_link_checkboxes.append(cb)

            link_section.set_content(link_content)
            link_layout.addWidget(link_section)
            setattr(self, "popup_c1_link_checkboxes", c1_link_checkboxes)

        subsections_target_layout.addWidget(link_container)

        # === 3. INSIGHT INTO TREATMENT NEED CONTAINER ===
        treatment_container = QFrame()
        treatment_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(22, 163, 74, 0.4);
                border-radius: 10px;
            }
        """)
        treatment_layout = QVBoxLayout(treatment_container)
        treatment_layout.setContentsMargins(12, 10, 12, 10)
        treatment_layout.setSpacing(8)

        treatment_lbl = QLabel("Insight into Need for Treatment")
        treatment_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        treatment_layout.addWidget(treatment_lbl)

        treatment_text = QTextEdit()
        treatment_text.setPlaceholderText("Describe attitude and behaviour toward treatment need...")
        treatment_text.setStyleSheet(input_style + " background: white;")
        treatment_text.setMinimumHeight(58)
        treatment_layout.addWidget(treatment_text)
        treatment_layout.addWidget(DragResizeBar(treatment_text))
        text_widgets.append(treatment_text)
        subsection_data.append(("Insight into Need for Treatment", treatment_text))
        setattr(self, "popup_c1_treatment", treatment_text)

        if CollapsibleSection:
            treatment_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            treatment_section.set_content_height(200)
            treatment_section._min_height = 80
            treatment_section._max_height = 300
            treatment_section.set_header_style("""
                QFrame {
                    background: rgba(22, 163, 74, 0.15);
                    border: 1px solid rgba(22, 163, 74, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            treatment_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #15803d;
                    background: transparent;
                    border: none;
                }
            """)

            treatment_content = QWidget()
            treatment_content.setStyleSheet("""
                QWidget {
                    background: rgba(240, 253, 244, 0.95);
                    border: 1px solid rgba(22, 163, 74, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            treatment_cb_layout = QVBoxLayout(treatment_content)
            treatment_cb_layout.setContentsMargins(10, 6, 10, 6)
            treatment_cb_layout.setSpacing(3)

            C1_TREATMENT_ITEMS = [
                ("tx_refuses_treatment", "Refuses treatment"),
                ("tx_non_concordant", "Non-concordant with medication"),
                ("tx_lacks_understanding", "Lacks understanding of need for treatment"),
                ("tx_only_under_compulsion", "Only accepts treatment under compulsion"),
                ("tx_recurrent_disengagement", "Recurrent disengagement from services"),
                ("tx_accepts_medication", "Accepts medication (protective)"),
                ("tx_engages_mdt", "Engages with MDT (protective)"),
                ("tx_requests_help", "Requests help when unwell (protective)"),
            ]

            c1_treatment_checkboxes = []
            for item_key, item_label in C1_TREATMENT_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c1_key", item_key)
                cb.setProperty("c1_category", "treatment")
                cb.stateChanged.connect(self._update_c1_insight_narrative)
                treatment_cb_layout.addWidget(cb)
                c1_treatment_checkboxes.append(cb)

            treatment_section.set_content(treatment_content)
            treatment_layout.addWidget(treatment_section)
            setattr(self, "popup_c1_treatment_checkboxes", c1_treatment_checkboxes)

        subsections_target_layout.addWidget(treatment_container)

        # === 4. STABILITY/FLUCTUATION OF INSIGHT CONTAINER ===
        stability_container = QFrame()
        stability_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 249, 195, 0.5);
                border: 2px solid rgba(202, 138, 4, 0.4);
                border-radius: 10px;
            }
        """)
        stability_layout = QVBoxLayout(stability_container)
        stability_layout.setContentsMargins(12, 10, 12, 10)
        stability_layout.setSpacing(8)

        stability_lbl = QLabel("Stability/Fluctuation of Insight")
        stability_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #a16207; background: transparent; border: none;")
        stability_layout.addWidget(stability_lbl)

        stability_text = QTextEdit()
        stability_text.setPlaceholderText("Describe consistency of insight over time, especially during relapse...")
        stability_text.setStyleSheet(input_style + " background: white;")
        stability_text.setMinimumHeight(58)
        stability_layout.addWidget(stability_text)
        stability_layout.addWidget(DragResizeBar(stability_text))
        text_widgets.append(stability_text)
        subsection_data.append(("Stability/Fluctuation of Insight", stability_text))
        setattr(self, "popup_c1_stability", stability_text)

        if CollapsibleSection:
            stability_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            stability_section.set_content_height(160)
            stability_section._min_height = 80
            stability_section._max_height = 250
            stability_section.set_header_style("""
                QFrame {
                    background: rgba(202, 138, 4, 0.15);
                    border: 1px solid rgba(202, 138, 4, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            stability_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #a16207;
                    background: transparent;
                    border: none;
                }
            """)

            stability_content = QWidget()
            stability_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(202, 138, 4, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            stability_cb_layout = QVBoxLayout(stability_content)
            stability_cb_layout.setContentsMargins(10, 6, 10, 6)
            stability_cb_layout.setSpacing(3)

            C1_STABILITY_ITEMS = [
                ("stab_fluctuates", "Insight fluctuates"),
                ("stab_improves_meds", "Insight improves with medication"),
                ("stab_poor_when_unwell", "Poor insight when acutely unwell"),
                ("stab_only_when_well", "Insight only present when well"),
                ("stab_lost_relapse", "Insight lost during relapse"),
            ]

            c1_stability_checkboxes = []
            for item_key, item_label in C1_STABILITY_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c1_key", item_key)
                cb.setProperty("c1_category", "stability")
                cb.stateChanged.connect(self._update_c1_insight_narrative)
                stability_cb_layout.addWidget(cb)
                c1_stability_checkboxes.append(cb)

            stability_section.set_content(stability_content)
            stability_layout.addWidget(stability_section)
            setattr(self, "popup_c1_stability_checkboxes", c1_stability_checkboxes)

        subsections_target_layout.addWidget(stability_container)

        # === 5. BEHAVIOURAL INDICATORS CONTAINER ===
        behaviour_container = QFrame()
        behaviour_container.setStyleSheet("""
            QFrame {
                background: rgba(243, 232, 255, 0.5);
                border: 2px solid rgba(147, 51, 234, 0.4);
                border-radius: 10px;
            }
        """)
        behaviour_layout = QVBoxLayout(behaviour_container)
        behaviour_layout.setContentsMargins(12, 10, 12, 10)
        behaviour_layout.setSpacing(8)

        behaviour_lbl = QLabel("Behavioural Indicators")
        behaviour_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #7e22ce; background: transparent; border: none;")
        behaviour_layout.addWidget(behaviour_lbl)

        behaviour_text = QTextEdit()
        behaviour_text.setPlaceholderText("Describe behavioural evidence of insight (actions vs words)...")
        behaviour_text.setStyleSheet(input_style + " background: white;")
        behaviour_text.setMinimumHeight(58)
        behaviour_layout.addWidget(behaviour_text)
        behaviour_layout.addWidget(DragResizeBar(behaviour_text))
        text_widgets.append(behaviour_text)
        subsection_data.append(("Behavioural Indicators", behaviour_text))
        setattr(self, "popup_c1_behaviour", behaviour_text)

        if CollapsibleSection:
            behaviour_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            behaviour_section.set_content_height(180)
            behaviour_section._min_height = 80
            behaviour_section._max_height = 280
            behaviour_section.set_header_style("""
                QFrame {
                    background: rgba(147, 51, 234, 0.15);
                    border: 1px solid rgba(147, 51, 234, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            behaviour_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #7e22ce;
                    background: transparent;
                    border: none;
                }
            """)

            behaviour_content = QWidget()
            behaviour_content.setStyleSheet("""
                QWidget {
                    background: rgba(250, 245, 255, 0.95);
                    border: 1px solid rgba(147, 51, 234, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            behaviour_cb_layout = QVBoxLayout(behaviour_content)
            behaviour_cb_layout.setContentsMargins(10, 6, 10, 6)
            behaviour_cb_layout.setSpacing(3)

            C1_BEHAVIOUR_ITEMS = [
                ("beh_stops_meds", "Stops medication after discharge"),
                ("beh_misses_appts", "Misses appointments"),
                ("beh_rejects_followup", "Rejects follow-up"),
                ("beh_blames_services", "Repeatedly blames services"),
                ("beh_recurrent_relapse", "Recurrent relapse after disengagement"),
                ("beh_consistent_engagement", "Consistent engagement (protective)"),
            ]

            c1_behaviour_checkboxes = []
            for item_key, item_label in C1_BEHAVIOUR_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c1_key", item_key)
                cb.setProperty("c1_category", "behaviour")
                cb.stateChanged.connect(self._update_c1_insight_narrative)
                behaviour_cb_layout.addWidget(cb)
                c1_behaviour_checkboxes.append(cb)

            behaviour_section.set_content(behaviour_content)
            behaviour_layout.addWidget(behaviour_section)
            setattr(self, "popup_c1_behaviour_checkboxes", c1_behaviour_checkboxes)

        subsections_target_layout.addWidget(behaviour_container)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")

            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_c1_insight_narrative(self):
        """Update C1 text fields with gender-sensitive narrative based on checked indicators."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()

        # Define protective keys for each category
        protective_keys = {
            "disorder": ["dis_accepts_diagnosis", "dis_recognises_symptoms"],
            "link": ["link_understands_triggers", "link_acknowledges_unwell"],
            "treatment": ["tx_accepts_medication", "tx_engages_mdt", "tx_requests_help"],
            "stability": [],  # All stability items are concerns
            "behaviour": ["beh_consistent_engagement"],
        }

        # Category configs: (checkboxes_attr, text_attr, category_key, narratives_dict)
        categories = [
            ("popup_c1_disorder_checkboxes", "popup_c1_disorder", "disorder", {
                "dis_denies_illness": f"{Subj} denies having any mental illness",
                "dis_rejects_diagnosis": f"{Subj} rejects {poss} diagnosis",
                "dis_external_attribution": f"{Subj} attributes {poss} symptoms to external factors",
                "dis_poor_insight": f"{Subj} has poor insight into {poss} mental disorder",
                "dis_no_recognise_relapse": f"{Subj} does not recognise early warning signs of relapse",
                "dis_accepts_diagnosis": f"{Subj} accepts {poss} diagnosis",
                "dis_recognises_symptoms": f"{Subj} recognises that {poss} symptoms are illness-related",
            }),
            ("popup_c1_link_checkboxes", "popup_c1_link", "link", {
                "link_denies_connection": f"{Subj} denies any link between {poss} illness and past offending",
                "link_minimises_violence": f"{Subj} minimises {poss} past violence",
                "link_externalises_blame": f"{Subj} externalises blame for {poss} violent behaviour",
                "link_lacks_victim_empathy": f"{Subj} lacks empathy for {poss} victims",
                "link_no_reflection": f"{Subj} shows limited reflection on the index offence",
                "link_understands_triggers": f"{Subj} understands the triggers for {poss} risk behaviour",
                "link_acknowledges_unwell": f"{Subj} acknowledges that {subj} was unwell at the time of the offence",
            }),
            ("popup_c1_treatment_checkboxes", "popup_c1_treatment", "treatment", {
                "tx_refuses_treatment": f"{Subj} refuses treatment",
                "tx_non_concordant": f"{Subj} is non-concordant with prescribed medication",
                "tx_lacks_understanding": f"{Subj} lacks understanding of {poss} need for treatment",
                "tx_only_under_compulsion": f"{Subj} only accepts treatment under compulsion",
                "tx_recurrent_disengagement": f"{Subj} has a pattern of recurrent disengagement from services",
                "tx_accepts_medication": f"{Subj} accepts {poss} medication",
                "tx_engages_mdt": f"{Subj} engages appropriately with the multidisciplinary team",
                "tx_requests_help": f"{Subj} requests help when {subj} feels unwell",
            }),
            ("popup_c1_stability_checkboxes", "popup_c1_stability", "stability", {
                "stab_fluctuates": f"{poss.capitalize()} insight fluctuates over time",
                "stab_improves_meds": f"{poss.capitalize()} insight improves with medication",
                "stab_poor_when_unwell": f"{Subj} has poor insight when acutely unwell",
                "stab_only_when_well": f"{poss.capitalize()} insight is only present when {subj} is well",
                "stab_lost_relapse": f"{poss.capitalize()} insight is lost during periods of relapse",
            }),
            ("popup_c1_behaviour_checkboxes", "popup_c1_behaviour", "behaviour", {
                "beh_stops_meds": f"{Subj} stops medication after discharge",
                "beh_misses_appts": f"{Subj} misses appointments",
                "beh_rejects_followup": f"{Subj} rejects follow-up care",
                "beh_blames_services": f"{Subj} repeatedly blames services for {poss} difficulties",
                "beh_recurrent_relapse": f"{Subj} has had recurrent relapses following disengagement",
                "beh_consistent_engagement": f"{Subj} demonstrates consistent engagement with services",
            }),
        ]

        def build_phrase_list(phrases):
            """Build a comma-separated phrase list with 'and' before last item."""
            if not phrases:
                return ""
            if len(phrases) == 1:
                return phrases[0]
            elif len(phrases) == 2:
                return phrases[0] + " and " + phrases[1][0].lower() + phrases[1][1:]
            else:
                result = phrases[0]
                for phrase in phrases[1:-1]:
                    result += ", " + phrase[0].lower() + phrase[1:]
                result += ", and " + phrases[-1][0].lower() + phrases[-1][1:]
                return result

        def build_stability_narrative(checked_keys):
            """Build custom flowing narrative for stability section avoiding repetition."""
            # Items: fluctuates(1), improves_meds(2-positive), poor_when_unwell(3),
            #        only_when_well(4-positive), lost_relapse(5)
            has_fluctuates = "stab_fluctuates" in checked_keys
            has_improves = "stab_improves_meds" in checked_keys
            has_poor = "stab_poor_when_unwell" in checked_keys
            has_only_well = "stab_only_when_well" in checked_keys
            has_lost = "stab_lost_relapse" in checked_keys

            if not any([has_fluctuates, has_improves, has_poor, has_only_well, has_lost]):
                return ""

            Poss = poss.capitalize()
            sentences = []

            # Build first sentence starting with "His/Her insight..."
            first_parts = []

            if has_fluctuates:
                first_parts.append("fluctuates over time")

            if has_improves:
                first_parts.append("improves with medication")

            # Build the first sentence
            if first_parts:
                first_sentence = f"{Poss} insight " + " and ".join(first_parts) if len(first_parts) <= 2 else f"{Poss} insight " + ", ".join(first_parts[:-1]) + " and " + first_parts[-1]

                # Add poor_when_unwell with "but" if we have positive items, otherwise continue
                if has_poor:
                    if has_improves:
                        first_sentence += f" but {subj} has poor insight when acutely unwell"
                    else:
                        first_sentence += f", it is poor when acutely unwell"

                # Add only_when_well
                if has_only_well:
                    if has_poor:
                        first_sentence += " and only present when " + subj + " is well"
                    elif has_improves or has_fluctuates:
                        first_sentence += " and is only present when " + subj + " is well"

                sentences.append(first_sentence + ".")
            else:
                # No fluctuates or improves - start differently
                if has_poor and has_only_well:
                    sentences.append(f"{Poss} insight is poor when acutely unwell and only present when {subj} is well.")
                elif has_poor:
                    sentences.append(f"{Subj} has poor insight when acutely unwell.")
                elif has_only_well:
                    sentences.append(f"{Poss} insight is only present when {subj} is well.")

            # Handle lost_relapse as separate sentence
            if has_lost:
                if sentences:
                    sentences.append(f"{Subj} loses insight during periods of relapse.")
                else:
                    sentences.append(f"{Poss} insight is lost during periods of relapse.")

            return " ".join(sentences)

        for cb_attr, text_attr, cat_key, narratives in categories:
            checkboxes = getattr(self, cb_attr, [])
            text_widget = getattr(self, text_attr, None)
            prot_keys = protective_keys.get(cat_key, [])

            # Special handling for stability section
            if cat_key == "stability":
                checked_keys = []
                for cb in checkboxes:
                    if cb.isChecked():
                        key = cb.property("c1_key")
                        if key:
                            checked_keys.append(key)
                if text_widget:
                    if checked_keys:
                        narrative = build_stability_narrative(checked_keys)
                        text_widget.setPlainText(narrative)
                    else:
                        text_widget.clear()
                continue

            vulnerability_phrases = []
            protective_phrases = []

            for cb in checkboxes:
                if cb.isChecked():
                    key = cb.property("c1_key")
                    if key and key in narratives:
                        if key in prot_keys:
                            protective_phrases.append(narratives[key])
                        else:
                            vulnerability_phrases.append(narratives[key])

            if text_widget:
                if vulnerability_phrases or protective_phrases:
                    sentences = []

                    # Build vulnerability sentences - split if more than 4
                    if vulnerability_phrases:
                        if len(vulnerability_phrases) <= 4:
                            vuln_text = build_phrase_list(vulnerability_phrases)
                            sentences.append(vuln_text + ".")
                        else:
                            # First 4 vulnerabilities in first sentence
                            first_part = build_phrase_list(vulnerability_phrases[:4])
                            sentences.append(first_part + ".")
                            # Remaining vulnerabilities with "He/She also"
                            remaining = vulnerability_phrases[4:]
                            if len(remaining) == 1:
                                sentences.append(f"{Subj} also " + remaining[0].replace(Subj + " ", "").replace(poss.capitalize() + " ", poss + " ") + ".")
                            else:
                                also_text = build_phrase_list(remaining)
                                sentences.append(f"{Subj} also " + also_text.replace(Subj + " ", "").replace(poss.capitalize() + " ", poss + " ") + ".")

                    # Build protective sentence with appropriate prefix
                    if protective_phrases:
                        prot_text = build_phrase_list(protective_phrases)
                        if vulnerability_phrases:
                            # Use "Nevertheless" only if vulnerabilities exist
                            sentences.append("Nevertheless, " + prot_text[0].lower() + prot_text[1:] + ".")
                        else:
                            sentences.append(prot_text + ".")

                    text_widget.setPlainText(" ".join(sentences))
                else:
                    text_widget.clear()

    def _build_hcr_c2_popup(self, key: str, code: str, description: str, subsections: list):
        """Build C2 (Violent Ideation or Intent) popup with tick boxes.

        Based on HCR-20V3 manual guidance for C2 - thoughts, fantasies, and
        intentions of harming others.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. EXPLICIT VIOLENT THOUGHTS CONTAINER ===
        explicit_container = QFrame()
        explicit_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        explicit_layout = QVBoxLayout(explicit_container)
        explicit_layout.setContentsMargins(12, 10, 12, 10)
        explicit_layout.setSpacing(8)

        explicit_lbl = QLabel("Explicit Violent Thoughts")
        explicit_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        explicit_layout.addWidget(explicit_lbl)

        explicit_text = QTextEdit()
        explicit_text.setPlaceholderText("Describe explicit violent thoughts, fantasies, or intent to harm...")
        explicit_text.setStyleSheet(input_style + " background: white;")
        explicit_text.setMinimumHeight(58)
        explicit_layout.addWidget(explicit_text)
        explicit_layout.addWidget(DragResizeBar(explicit_text))
        text_widgets.append(explicit_text)
        subsection_data.append(("Explicit Violent Thoughts", explicit_text))
        setattr(self, "popup_c2_explicit", explicit_text)

        if CollapsibleSection:
            explicit_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            explicit_section.set_content_height(180)
            explicit_section._min_height = 80
            explicit_section._max_height = 280
            explicit_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            explicit_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            explicit_content = QWidget()
            explicit_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            explicit_cb_layout = QVBoxLayout(explicit_content)
            explicit_cb_layout.setContentsMargins(10, 6, 10, 6)
            explicit_cb_layout.setSpacing(3)

            C2_EXPLICIT_ITEMS = [
                ("exp_thoughts_harm", "Thoughts of harming others"),
                ("exp_violent_thoughts", "Violent thoughts documented"),
                ("exp_homicidal_ideation", "Homicidal ideation"),
                ("exp_desire_assault", "Desire to assault someone"),
                ("exp_kill_fantasies", "Fantasies about killing"),
                ("exp_specific_target", "Specific target identified"),
            ]

            c2_explicit_checkboxes = []
            for item_key, item_label in C2_EXPLICIT_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c2_key", item_key)
                cb.setProperty("c2_category", "explicit")
                cb.stateChanged.connect(self._update_c2_ideation_narrative)
                explicit_cb_layout.addWidget(cb)
                c2_explicit_checkboxes.append(cb)

            explicit_section.set_content(explicit_content)
            explicit_layout.addWidget(explicit_section)
            setattr(self, "popup_c2_explicit_checkboxes", c2_explicit_checkboxes)

        subsections_target_layout.addWidget(explicit_container)

        # === 2. CONDITIONAL VIOLENCE CONTAINER ===
        conditional_container = QFrame()
        conditional_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        conditional_layout = QVBoxLayout(conditional_container)
        conditional_layout.setContentsMargins(12, 10, 12, 10)
        conditional_layout.setSpacing(8)

        conditional_lbl = QLabel("Conditional Violence")
        conditional_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #b45309; background: transparent; border: none;")
        conditional_layout.addWidget(conditional_lbl)

        conditional_text = QTextEdit()
        conditional_text.setPlaceholderText("Describe conditional or contingent threats (e.g., 'if X happens, I'll...')...")
        conditional_text.setStyleSheet(input_style + " background: white;")
        conditional_text.setMinimumHeight(58)
        conditional_layout.addWidget(conditional_text)
        conditional_layout.addWidget(DragResizeBar(conditional_text))
        text_widgets.append(conditional_text)
        subsection_data.append(("Conditional Violence", conditional_text))
        setattr(self, "popup_c2_conditional", conditional_text)

        if CollapsibleSection:
            conditional_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            conditional_section.set_content_height(160)
            conditional_section._min_height = 80
            conditional_section._max_height = 250
            conditional_section.set_header_style("""
                QFrame {
                    background: rgba(217, 119, 6, 0.15);
                    border: 1px solid rgba(217, 119, 6, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            conditional_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #b45309;
                    background: transparent;
                    border: none;
                }
            """)

            conditional_content = QWidget()
            conditional_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(217, 119, 6, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            conditional_cb_layout = QVBoxLayout(conditional_content)
            conditional_cb_layout.setContentsMargins(10, 6, 10, 6)
            conditional_cb_layout.setSpacing(3)

            C2_CONDITIONAL_ITEMS = [
                ("cond_if_provoked", "States will act 'if provoked'"),
                ("cond_self_defence", "Claims violence for 'self-defence'"),
                ("cond_snap", "Warns may 'snap' or lose control"),
                ("cond_someone_hurt", "States 'someone will get hurt'"),
                ("cond_dont_know", "Says 'don't know what I'll do'"),
            ]

            c2_conditional_checkboxes = []
            for item_key, item_label in C2_CONDITIONAL_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c2_key", item_key)
                cb.setProperty("c2_category", "conditional")
                cb.stateChanged.connect(self._update_c2_ideation_narrative)
                conditional_cb_layout.addWidget(cb)
                c2_conditional_checkboxes.append(cb)

            conditional_section.set_content(conditional_content)
            conditional_layout.addWidget(conditional_section)
            setattr(self, "popup_c2_conditional_checkboxes", c2_conditional_checkboxes)

        subsections_target_layout.addWidget(conditional_container)

        # === 3. JUSTIFICATION/ENDORSEMENT CONTAINER ===
        justify_container = QFrame()
        justify_container.setStyleSheet("""
            QFrame {
                background: rgba(243, 232, 255, 0.5);
                border: 2px solid rgba(147, 51, 234, 0.4);
                border-radius: 10px;
            }
        """)
        justify_layout = QVBoxLayout(justify_container)
        justify_layout.setContentsMargins(12, 10, 12, 10)
        justify_layout.setSpacing(8)

        justify_lbl = QLabel("Justification/Endorsement of Violence")
        justify_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #7e22ce; background: transparent; border: none;")
        justify_layout.addWidget(justify_lbl)

        justify_text = QTextEdit()
        justify_text.setPlaceholderText("Describe justification, minimisation, or endorsement of violence...")
        justify_text.setStyleSheet(input_style + " background: white;")
        justify_text.setMinimumHeight(58)
        justify_layout.addWidget(justify_text)
        justify_layout.addWidget(DragResizeBar(justify_text))
        text_widgets.append(justify_text)
        subsection_data.append(("Justification/Endorsement", justify_text))
        setattr(self, "popup_c2_justify", justify_text)

        if CollapsibleSection:
            justify_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            justify_section.set_content_height(160)
            justify_section._min_height = 80
            justify_section._max_height = 250
            justify_section.set_header_style("""
                QFrame {
                    background: rgba(147, 51, 234, 0.15);
                    border: 1px solid rgba(147, 51, 234, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            justify_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #7e22ce;
                    background: transparent;
                    border: none;
                }
            """)

            justify_content = QWidget()
            justify_content.setStyleSheet("""
                QWidget {
                    background: rgba(250, 245, 255, 0.95);
                    border: 1px solid rgba(147, 51, 234, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            justify_cb_layout = QVBoxLayout(justify_content)
            justify_cb_layout.setContentsMargins(10, 6, 10, 6)
            justify_cb_layout.setSpacing(3)

            C2_JUSTIFY_ITEMS = [
                ("just_deserved_it", "Believes victim 'deserved it'"),
                ("just_provoked", "Claims was provoked"),
                ("just_no_choice", "Says 'had no choice'"),
                ("just_anyone_same", "States 'anyone would have done the same'"),
                ("just_necessary", "Views violence as necessary response"),
            ]

            c2_justify_checkboxes = []
            for item_key, item_label in C2_JUSTIFY_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c2_key", item_key)
                cb.setProperty("c2_category", "justify")
                cb.stateChanged.connect(self._update_c2_ideation_narrative)
                justify_cb_layout.addWidget(cb)
                c2_justify_checkboxes.append(cb)

            justify_section.set_content(justify_content)
            justify_layout.addWidget(justify_section)
            setattr(self, "popup_c2_justify_checkboxes", c2_justify_checkboxes)

        subsections_target_layout.addWidget(justify_container)

        # === 4. IDEATION LINKED TO SYMPTOMS CONTAINER ===
        symptoms_container = QFrame()
        symptoms_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        symptoms_layout = QVBoxLayout(symptoms_container)
        symptoms_layout.setContentsMargins(12, 10, 12, 10)
        symptoms_layout.setSpacing(8)

        symptoms_lbl = QLabel("Ideation Linked to Mental State")
        symptoms_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        symptoms_layout.addWidget(symptoms_lbl)

        symptoms_text = QTextEdit()
        symptoms_text.setPlaceholderText("Describe violent ideation linked to psychosis, delusions, or command hallucinations...")
        symptoms_text.setStyleSheet(input_style + " background: white;")
        symptoms_text.setMinimumHeight(58)
        symptoms_layout.addWidget(symptoms_text)
        symptoms_layout.addWidget(DragResizeBar(symptoms_text))
        text_widgets.append(symptoms_text)
        subsection_data.append(("Ideation Linked to Mental State", symptoms_text))
        setattr(self, "popup_c2_symptoms", symptoms_text)

        if CollapsibleSection:
            symptoms_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            symptoms_section.set_content_height(160)
            symptoms_section._min_height = 80
            symptoms_section._max_height = 250
            symptoms_section.set_header_style("""
                QFrame {
                    background: rgba(59, 130, 246, 0.15);
                    border: 1px solid rgba(59, 130, 246, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            symptoms_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #1d4ed8;
                    background: transparent;
                    border: none;
                }
            """)

            symptoms_content = QWidget()
            symptoms_content.setStyleSheet("""
                QWidget {
                    background: rgba(239, 246, 255, 0.95);
                    border: 1px solid rgba(59, 130, 246, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            symptoms_cb_layout = QVBoxLayout(symptoms_content)
            symptoms_cb_layout.setContentsMargins(10, 6, 10, 6)
            symptoms_cb_layout.setSpacing(3)

            C2_SYMPTOMS_ITEMS = [
                ("sym_command_hallucinations", "Command hallucinations to harm"),
                ("sym_voices_harm", "Voices telling to hurt others"),
                ("sym_paranoid_retaliation", "Paranoid with retaliatory intent"),
                ("sym_psychotic_violent", "Violent thoughts when psychotic"),
                ("sym_persecutory_beliefs", "Persecutory beliefs with violent response"),
            ]

            c2_symptoms_checkboxes = []
            for item_key, item_label in C2_SYMPTOMS_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c2_key", item_key)
                cb.setProperty("c2_category", "symptoms")
                cb.stateChanged.connect(self._update_c2_ideation_narrative)
                symptoms_cb_layout.addWidget(cb)
                c2_symptoms_checkboxes.append(cb)

            symptoms_section.set_content(symptoms_content)
            symptoms_layout.addWidget(symptoms_section)
            setattr(self, "popup_c2_symptoms_checkboxes", c2_symptoms_checkboxes)

        subsections_target_layout.addWidget(symptoms_container)

        # === 5. AGGRESSIVE RUMINATION CONTAINER ===
        rumination_container = QFrame()
        rumination_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 249, 195, 0.5);
                border: 2px solid rgba(202, 138, 4, 0.4);
                border-radius: 10px;
            }
        """)
        rumination_layout = QVBoxLayout(rumination_container)
        rumination_layout.setContentsMargins(12, 10, 12, 10)
        rumination_layout.setSpacing(8)

        rumination_lbl = QLabel("Aggressive Rumination")
        rumination_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #a16207; background: transparent; border: none;")
        rumination_layout.addWidget(rumination_lbl)

        rumination_text = QTextEdit()
        rumination_text.setPlaceholderText("Describe brooding, grievance, persistent anger, or revenge thoughts...")
        rumination_text.setStyleSheet(input_style + " background: white;")
        rumination_text.setMinimumHeight(58)
        rumination_layout.addWidget(rumination_text)
        rumination_layout.addWidget(DragResizeBar(rumination_text))
        text_widgets.append(rumination_text)
        subsection_data.append(("Aggressive Rumination", rumination_text))
        setattr(self, "popup_c2_rumination", rumination_text)

        if CollapsibleSection:
            rumination_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            rumination_section.set_content_height(160)
            rumination_section._min_height = 80
            rumination_section._max_height = 250
            rumination_section.set_header_style("""
                QFrame {
                    background: rgba(202, 138, 4, 0.15);
                    border: 1px solid rgba(202, 138, 4, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            rumination_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #a16207;
                    background: transparent;
                    border: none;
                }
            """)

            rumination_content = QWidget()
            rumination_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(202, 138, 4, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            rumination_cb_layout = QVBoxLayout(rumination_content)
            rumination_cb_layout.setContentsMargins(10, 6, 10, 6)
            rumination_cb_layout.setSpacing(3)

            C2_RUMINATION_ITEMS = [
                ("rum_persistent_anger", "Persistent anger documented"),
                ("rum_grievance", "Holds grievances"),
                ("rum_grudges", "Holds grudges"),
                ("rum_brooding", "Brooding behaviour"),
                ("rum_revenge", "Thoughts of revenge"),
                ("rum_escalating", "Escalating language/preoccupation"),
            ]

            c2_rumination_checkboxes = []
            for item_key, item_label in C2_RUMINATION_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c2_key", item_key)
                cb.setProperty("c2_category", "rumination")
                cb.stateChanged.connect(self._update_c2_ideation_narrative)
                rumination_cb_layout.addWidget(cb)
                c2_rumination_checkboxes.append(cb)

            rumination_section.set_content(rumination_content)
            rumination_layout.addWidget(rumination_section)
            setattr(self, "popup_c2_rumination_checkboxes", c2_rumination_checkboxes)

        subsections_target_layout.addWidget(rumination_container)

        # === 6. THREATS CONTAINER ===
        threats_container = QFrame()
        threats_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(22, 163, 74, 0.4);
                border-radius: 10px;
            }
        """)
        threats_layout = QVBoxLayout(threats_container)
        threats_layout.setContentsMargins(12, 10, 12, 10)
        threats_layout.setSpacing(8)

        threats_lbl = QLabel("Threats")
        threats_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        threats_layout.addWidget(threats_lbl)

        threats_text = QTextEdit()
        threats_text.setPlaceholderText("Describe verbal threats, intimidation, or threatening behaviour...")
        threats_text.setStyleSheet(input_style + " background: white;")
        threats_text.setMinimumHeight(58)
        threats_layout.addWidget(threats_text)
        threats_layout.addWidget(DragResizeBar(threats_text))
        text_widgets.append(threats_text)
        subsection_data.append(("Threats", threats_text))
        setattr(self, "popup_c2_threats", threats_text)

        if CollapsibleSection:
            threats_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            threats_section.set_content_height(160)
            threats_section._min_height = 80
            threats_section._max_height = 250
            threats_section.set_header_style("""
                QFrame {
                    background: rgba(22, 163, 74, 0.15);
                    border: 1px solid rgba(22, 163, 74, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            threats_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #15803d;
                    background: transparent;
                    border: none;
                }
            """)

            threats_content = QWidget()
            threats_content.setStyleSheet("""
                QWidget {
                    background: rgba(240, 253, 244, 0.95);
                    border: 1px solid rgba(22, 163, 74, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            threats_cb_layout = QVBoxLayout(threats_content)
            threats_cb_layout.setContentsMargins(10, 6, 10, 6)
            threats_cb_layout.setSpacing(3)

            C2_THREATS_ITEMS = [
                ("thr_verbal_threats", "Verbal threats made"),
                ("thr_threatened_staff", "Threatened staff"),
                ("thr_threatened_family", "Threatened family members"),
                ("thr_intimidating", "Intimidating behaviour"),
                ("thr_aggressive_statements", "Aggressive statements"),
                ("thr_no_follow_through", "Threats but no follow-through (note pattern)"),
            ]

            c2_threats_checkboxes = []
            for item_key, item_label in C2_THREATS_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c2_key", item_key)
                cb.setProperty("c2_category", "threats")
                cb.stateChanged.connect(self._update_c2_ideation_narrative)
                threats_cb_layout.addWidget(cb)
                c2_threats_checkboxes.append(cb)

            threats_section.set_content(threats_content)
            threats_layout.addWidget(threats_section)
            setattr(self, "popup_c2_threats_checkboxes", c2_threats_checkboxes)

        subsections_target_layout.addWidget(threats_container)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")

            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_c2_ideation_narrative(self):
        """Update C2 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. EXPLICIT IDEATION ============
        exp_items = {
            "exp_thoughts_harm": "thoughts of harming others",
            "exp_violent_thoughts": "documented violent thoughts",
            "exp_homicidal_ideation": "homicidal ideation",
            "exp_desire_assault": "a desire to assault someone",
            "exp_kill_fantasies": "fantasies about killing",
            "exp_specific_target": "a specific target identified",
        }

        exp_cbs = getattr(self, "popup_c2_explicit_checkboxes", [])
        exp_text = getattr(self, "popup_c2_explicit", None)

        exp_checked = []
        for cb in exp_cbs:
            if cb.isChecked():
                key = cb.property("c2_key")
                if key in exp_items:
                    exp_checked.append(exp_items[key])

        if exp_text:
            if exp_checked:
                exp_text.setPlainText(f"{Subj} has expressed explicit violent ideation, including {join_items(exp_checked)}.")
            else:
                exp_text.clear()

        # ============ 2. CONDITIONAL VIOLENCE ============
        cond_items = {
            "cond_if_provoked": "stating {subj} will act if provoked",
            "cond_self_defence": "claiming any violence would be self-defence",
            "cond_snap": "warning that {subj} may snap or lose control",
            "cond_someone_hurt": "stating that someone will get hurt",
            "cond_dont_know": "saying {subj} doesn't know what {subj}'ll do",
        }

        cond_cbs = getattr(self, "popup_c2_conditional_checkboxes", [])
        cond_text = getattr(self, "popup_c2_conditional", None)

        cond_checked = []
        for cb in cond_cbs:
            if cb.isChecked():
                key = cb.property("c2_key")
                if key in cond_items:
                    cond_checked.append(cond_items[key].format(subj=subj))

        if cond_text:
            if cond_checked:
                cond_text.setPlainText(f"{Subj} has made conditional statements about violence, {join_items(cond_checked)}.")
            else:
                cond_text.clear()

        # ============ 3. JUSTIFICATION/ENDORSEMENT ============
        just_items = {
            "just_deserved_it": "believing {poss} victim deserved it",
            "just_provoked": "claiming {subj} was provoked",
            "just_no_choice": "saying {subj} had no choice",
            "just_anyone_same": "stating anyone would have done the same",
            "just_necessary": "viewing violence as a necessary response",
        }

        just_cbs = getattr(self, "popup_c2_justify_checkboxes", [])
        just_text = getattr(self, "popup_c2_justify", None)

        just_checked = []
        for cb in just_cbs:
            if cb.isChecked():
                key = cb.property("c2_key")
                if key in just_items:
                    just_checked.append(just_items[key].format(poss=poss, subj=subj))

        if just_text:
            if just_checked:
                just_text.setPlainText(f"{Subj} demonstrates justification and endorsement of violence, {join_items(just_checked)}.")
            else:
                just_text.clear()

        # ============ 4. IDEATION LINKED TO MENTAL STATE ============
        sym_items = {
            "sym_command_hallucinations": "command hallucinations to harm",
            "sym_voices_harm": "voices telling {obj} to hurt others",
            "sym_paranoid_retaliation": "paranoid beliefs with retaliatory intent",
            "sym_psychotic_violent": "violent thoughts when psychotic",
            "sym_persecutory_beliefs": "persecutory beliefs that may lead to violent response",
        }

        sym_cbs = getattr(self, "popup_c2_symptoms_checkboxes", [])
        sym_text = getattr(self, "popup_c2_symptoms", None)

        sym_checked = []
        for cb in sym_cbs:
            if cb.isChecked():
                key = cb.property("c2_key")
                if key in sym_items:
                    sym_checked.append(sym_items[key].format(obj=obj))

        if sym_text:
            if sym_checked:
                sym_text.setPlainText(f"Violent ideation is linked to {poss} mental state, including {join_items(sym_checked)}.")
            else:
                sym_text.clear()

        # ============ 5. AGGRESSIVE RUMINATION ============
        rum_items = {
            "rum_persistent_anger": "persistent anger",
            "rum_grievance": "holding grievances",
            "rum_grudges": "holding grudges",
            "rum_brooding": "brooding behaviour",
            "rum_revenge": "thoughts of revenge",
            "rum_escalating": "escalating language and preoccupation",
        }

        rum_cbs = getattr(self, "popup_c2_rumination_checkboxes", [])
        rum_text = getattr(self, "popup_c2_rumination", None)

        rum_checked = []
        for cb in rum_cbs:
            if cb.isChecked():
                key = cb.property("c2_key")
                if key in rum_items:
                    rum_checked.append(rum_items[key])

        if rum_text:
            if rum_checked:
                rum_text.setPlainText(f"{Subj} demonstrates aggressive rumination, with {join_items(rum_checked)}.")
            else:
                rum_text.clear()

        # ============ 6. THREATS ============
        thr_items = {
            "thr_verbal_threats": "verbal threats",
            "thr_threatened_staff": "threatening staff",
            "thr_threatened_family": "threatening family members",
            "thr_intimidating": "intimidating behaviour",
            "thr_aggressive_statements": "aggressive statements",
            "thr_no_follow_through": "threats without follow-through",
        }

        thr_cbs = getattr(self, "popup_c2_threats_checkboxes", [])
        thr_text = getattr(self, "popup_c2_threats", None)

        thr_checked = []
        for cb in thr_cbs:
            if cb.isChecked():
                key = cb.property("c2_key")
                if key in thr_items:
                    thr_checked.append(thr_items[key])

        if thr_text:
            if thr_checked:
                thr_text.setPlainText(f"There is a documented history of threats, including {join_items(thr_checked)}.")
            else:
                thr_text.clear()

    def _build_hcr_c3_popup(self, key: str, code: str, description: str, subsections: list):
        """Build C3 (Symptoms of Major Mental Disorder) popup with tick boxes.

        Based on HCR-20V3 manual guidance for C3 - active symptoms of major
        mental disorder that are relevant to violence risk.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. PSYCHOTIC SYMPTOMS CONTAINER ===
        psychotic_container = QFrame()
        psychotic_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        psychotic_layout = QVBoxLayout(psychotic_container)
        psychotic_layout.setContentsMargins(12, 10, 12, 10)
        psychotic_layout.setSpacing(8)

        psychotic_lbl = QLabel("Psychotic Symptoms")
        psychotic_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        psychotic_layout.addWidget(psychotic_lbl)

        psychotic_text = QTextEdit()
        psychotic_text.setPlaceholderText("Describe psychotic symptoms (delusions, hallucinations, thought disorder)...")
        psychotic_text.setStyleSheet(input_style + " background: white;")
        psychotic_text.setMinimumHeight(58)
        psychotic_layout.addWidget(psychotic_text)
        psychotic_layout.addWidget(DragResizeBar(psychotic_text))
        text_widgets.append(psychotic_text)
        subsection_data.append(("Psychotic Symptoms", psychotic_text))
        setattr(self, "popup_c3_psychotic", psychotic_text)

        if CollapsibleSection:
            psychotic_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            psychotic_section.set_content_height(200)
            psychotic_section._min_height = 80
            psychotic_section._max_height = 300
            psychotic_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            psychotic_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            psychotic_content = QWidget()
            psychotic_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            psychotic_cb_layout = QVBoxLayout(psychotic_content)
            psychotic_cb_layout.setContentsMargins(10, 6, 10, 6)
            psychotic_cb_layout.setSpacing(3)

            C3_PSYCHOTIC_ITEMS = [
                ("psy_paranoid", "Paranoid ideation"),
                ("psy_persecutory", "Persecutory delusions"),
                ("psy_command_hallucinations", "Command hallucinations"),
                ("psy_hearing_voices", "Hearing voices"),
                ("psy_grandiose", "Grandiose delusions"),
                ("psy_thought_disorder", "Thought disorder"),
                ("psy_actively_psychotic", "Actively psychotic currently"),
            ]

            c3_psychotic_checkboxes = []
            for item_key, item_label in C3_PSYCHOTIC_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c3_key", item_key)
                cb.setProperty("c3_category", "psychotic")
                cb.stateChanged.connect(self._update_c3_symptoms_narrative)
                psychotic_cb_layout.addWidget(cb)
                c3_psychotic_checkboxes.append(cb)

            psychotic_section.set_content(psychotic_content)
            psychotic_layout.addWidget(psychotic_section)
            setattr(self, "popup_c3_psychotic_checkboxes", c3_psychotic_checkboxes)

        subsections_target_layout.addWidget(psychotic_container)

        # === 2. MANIA/HYPOMANIA CONTAINER ===
        mania_container = QFrame()
        mania_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        mania_layout = QVBoxLayout(mania_container)
        mania_layout.setContentsMargins(12, 10, 12, 10)
        mania_layout.setSpacing(8)

        mania_lbl = QLabel("Mania/Hypomania")
        mania_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #b45309; background: transparent; border: none;")
        mania_layout.addWidget(mania_lbl)

        mania_text = QTextEdit()
        mania_text.setPlaceholderText("Describe manic/hypomanic symptoms (elevated mood, grandiosity, disinhibition)...")
        mania_text.setStyleSheet(input_style + " background: white;")
        mania_text.setMinimumHeight(58)
        mania_layout.addWidget(mania_text)
        mania_layout.addWidget(DragResizeBar(mania_text))
        text_widgets.append(mania_text)
        subsection_data.append(("Mania/Hypomania", mania_text))
        setattr(self, "popup_c3_mania", mania_text)

        if CollapsibleSection:
            mania_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            mania_section.set_content_height(160)
            mania_section._min_height = 80
            mania_section._max_height = 250
            mania_section.set_header_style("""
                QFrame {
                    background: rgba(217, 119, 6, 0.15);
                    border: 1px solid rgba(217, 119, 6, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            mania_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #b45309;
                    background: transparent;
                    border: none;
                }
            """)

            mania_content = QWidget()
            mania_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(217, 119, 6, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            mania_cb_layout = QVBoxLayout(mania_content)
            mania_cb_layout.setContentsMargins(10, 6, 10, 6)
            mania_cb_layout.setSpacing(3)

            C3_MANIA_ITEMS = [
                ("man_manic", "Manic episode"),
                ("man_hypomanic", "Hypomanic episode"),
                ("man_elevated_mood", "Elevated mood"),
                ("man_grandiosity", "Grandiosity"),
                ("man_disinhibited", "Disinhibited behaviour"),
                ("man_reduced_sleep", "Reduced need for sleep"),
            ]

            c3_mania_checkboxes = []
            for item_key, item_label in C3_MANIA_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c3_key", item_key)
                cb.setProperty("c3_category", "mania")
                cb.stateChanged.connect(self._update_c3_symptoms_narrative)
                mania_cb_layout.addWidget(cb)
                c3_mania_checkboxes.append(cb)

            mania_section.set_content(mania_content)
            mania_layout.addWidget(mania_section)
            setattr(self, "popup_c3_mania_checkboxes", c3_mania_checkboxes)

        subsections_target_layout.addWidget(mania_container)

        # === 3. SEVERE DEPRESSION CONTAINER ===
        depression_container = QFrame()
        depression_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        depression_layout = QVBoxLayout(depression_container)
        depression_layout.setContentsMargins(12, 10, 12, 10)
        depression_layout.setSpacing(8)

        depression_lbl = QLabel("Severe Depression")
        depression_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        depression_layout.addWidget(depression_lbl)

        depression_text = QTextEdit()
        depression_text.setPlaceholderText("Describe severe depressive symptoms (agitation, hopelessness, nihilism)...")
        depression_text.setStyleSheet(input_style + " background: white;")
        depression_text.setMinimumHeight(58)
        depression_layout.addWidget(depression_text)
        depression_layout.addWidget(DragResizeBar(depression_text))
        text_widgets.append(depression_text)
        subsection_data.append(("Severe Depression", depression_text))
        setattr(self, "popup_c3_depression", depression_text)

        if CollapsibleSection:
            depression_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            depression_section.set_content_height(140)
            depression_section._min_height = 80
            depression_section._max_height = 220
            depression_section.set_header_style("""
                QFrame {
                    background: rgba(59, 130, 246, 0.15);
                    border: 1px solid rgba(59, 130, 246, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            depression_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #1d4ed8;
                    background: transparent;
                    border: none;
                }
            """)

            depression_content = QWidget()
            depression_content.setStyleSheet("""
                QWidget {
                    background: rgba(239, 246, 255, 0.95);
                    border: 1px solid rgba(59, 130, 246, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            depression_cb_layout = QVBoxLayout(depression_content)
            depression_cb_layout.setContentsMargins(10, 6, 10, 6)
            depression_cb_layout.setSpacing(3)

            C3_DEPRESSION_ITEMS = [
                ("dep_severe", "Severe depression"),
                ("dep_agitated", "Agitated depression"),
                ("dep_hopelessness", "Hopelessness with anger"),
                ("dep_nihilistic", "Nihilistic beliefs"),
                ("dep_paranoid", "Paranoid depression"),
            ]

            c3_depression_checkboxes = []
            for item_key, item_label in C3_DEPRESSION_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c3_key", item_key)
                cb.setProperty("c3_category", "depression")
                cb.stateChanged.connect(self._update_c3_symptoms_narrative)
                depression_cb_layout.addWidget(cb)
                c3_depression_checkboxes.append(cb)

            depression_section.set_content(depression_content)
            depression_layout.addWidget(depression_section)
            setattr(self, "popup_c3_depression_checkboxes", c3_depression_checkboxes)

        subsections_target_layout.addWidget(depression_container)

        # === 4. AFFECTIVE INSTABILITY CONTAINER ===
        affective_container = QFrame()
        affective_container.setStyleSheet("""
            QFrame {
                background: rgba(243, 232, 255, 0.5);
                border: 2px solid rgba(147, 51, 234, 0.4);
                border-radius: 10px;
            }
        """)
        affective_layout = QVBoxLayout(affective_container)
        affective_layout.setContentsMargins(12, 10, 12, 10)
        affective_layout.setSpacing(8)

        affective_lbl = QLabel("Affective Instability")
        affective_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #7e22ce; background: transparent; border: none;")
        affective_layout.addWidget(affective_lbl)

        affective_text = QTextEdit()
        affective_text.setPlaceholderText("Describe affective instability (labile mood, explosive anger, low frustration tolerance)...")
        affective_text.setStyleSheet(input_style + " background: white;")
        affective_text.setMinimumHeight(58)
        affective_layout.addWidget(affective_text)
        affective_layout.addWidget(DragResizeBar(affective_text))
        text_widgets.append(affective_text)
        subsection_data.append(("Affective Instability", affective_text))
        setattr(self, "popup_c3_affective", affective_text)

        if CollapsibleSection:
            affective_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            affective_section.set_content_height(140)
            affective_section._min_height = 80
            affective_section._max_height = 220
            affective_section.set_header_style("""
                QFrame {
                    background: rgba(147, 51, 234, 0.15);
                    border: 1px solid rgba(147, 51, 234, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            affective_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #7e22ce;
                    background: transparent;
                    border: none;
                }
            """)

            affective_content = QWidget()
            affective_content.setStyleSheet("""
                QWidget {
                    background: rgba(250, 245, 255, 0.95);
                    border: 1px solid rgba(147, 51, 234, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            affective_cb_layout = QVBoxLayout(affective_content)
            affective_cb_layout.setContentsMargins(10, 6, 10, 6)
            affective_cb_layout.setSpacing(3)

            C3_AFFECTIVE_ITEMS = [
                ("aff_labile", "Affect labile"),
                ("aff_easily_provoked", "Easily provoked"),
                ("aff_low_frustration", "Low frustration tolerance"),
                ("aff_explosive", "Explosive anger"),
                ("aff_rapid_shifts", "Rapid mood shifts"),
            ]

            c3_affective_checkboxes = []
            for item_key, item_label in C3_AFFECTIVE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c3_key", item_key)
                cb.setProperty("c3_category", "affective")
                cb.stateChanged.connect(self._update_c3_symptoms_narrative)
                affective_cb_layout.addWidget(cb)
                c3_affective_checkboxes.append(cb)

            affective_section.set_content(affective_content)
            affective_layout.addWidget(affective_section)
            setattr(self, "popup_c3_affective_checkboxes", c3_affective_checkboxes)

        subsections_target_layout.addWidget(affective_container)

        # === 5. AROUSAL/ANXIETY CONTAINER ===
        arousal_container = QFrame()
        arousal_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 249, 195, 0.5);
                border: 2px solid rgba(202, 138, 4, 0.4);
                border-radius: 10px;
            }
        """)
        arousal_layout = QVBoxLayout(arousal_container)
        arousal_layout.setContentsMargins(12, 10, 12, 10)
        arousal_layout.setSpacing(8)

        arousal_lbl = QLabel("Arousal/Anxiety States")
        arousal_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #a16207; background: transparent; border: none;")
        arousal_layout.addWidget(arousal_lbl)

        arousal_text = QTextEdit()
        arousal_text.setPlaceholderText("Describe arousal/anxiety symptoms (hypervigilance, heightened threat perception)...")
        arousal_text.setStyleSheet(input_style + " background: white;")
        arousal_text.setMinimumHeight(58)
        arousal_layout.addWidget(arousal_text)
        arousal_layout.addWidget(DragResizeBar(arousal_text))
        text_widgets.append(arousal_text)
        subsection_data.append(("Arousal/Anxiety States", arousal_text))
        setattr(self, "popup_c3_arousal", arousal_text)

        if CollapsibleSection:
            arousal_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            arousal_section.set_content_height(120)
            arousal_section._min_height = 80
            arousal_section._max_height = 200
            arousal_section.set_header_style("""
                QFrame {
                    background: rgba(202, 138, 4, 0.15);
                    border: 1px solid rgba(202, 138, 4, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            arousal_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #a16207;
                    background: transparent;
                    border: none;
                }
            """)

            arousal_content = QWidget()
            arousal_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(202, 138, 4, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            arousal_cb_layout = QVBoxLayout(arousal_content)
            arousal_cb_layout.setContentsMargins(10, 6, 10, 6)
            arousal_cb_layout.setSpacing(3)

            C3_AROUSAL_ITEMS = [
                ("ars_hypervigilant", "Hypervigilant"),
                ("ars_on_edge", "On edge / tense"),
                ("ars_threat_perception", "Heightened threat perception"),
                ("ars_ptsd_exacerbated", "PTSD symptoms exacerbated"),
            ]

            c3_arousal_checkboxes = []
            for item_key, item_label in C3_AROUSAL_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c3_key", item_key)
                cb.setProperty("c3_category", "arousal")
                cb.stateChanged.connect(self._update_c3_symptoms_narrative)
                arousal_cb_layout.addWidget(cb)
                c3_arousal_checkboxes.append(cb)

            arousal_section.set_content(arousal_content)
            arousal_layout.addWidget(arousal_section)
            setattr(self, "popup_c3_arousal_checkboxes", c3_arousal_checkboxes)

        subsections_target_layout.addWidget(arousal_container)

        # === 6. SYMPTOMS LINKED TO VIOLENCE CONTAINER ===
        linked_container = QFrame()
        linked_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(22, 163, 74, 0.4);
                border-radius: 10px;
            }
        """)
        linked_layout = QVBoxLayout(linked_container)
        linked_layout.setContentsMargins(12, 10, 12, 10)
        linked_layout.setSpacing(8)

        linked_lbl = QLabel("Symptoms Linked to Violence Risk")
        linked_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        linked_layout.addWidget(linked_lbl)

        linked_text = QTextEdit()
        linked_text.setPlaceholderText("Describe link between symptoms and past violence or current risk...")
        linked_text.setStyleSheet(input_style + " background: white;")
        linked_text.setMinimumHeight(58)
        linked_layout.addWidget(linked_text)
        linked_layout.addWidget(DragResizeBar(linked_text))
        text_widgets.append(linked_text)
        subsection_data.append(("Symptoms Linked to Violence Risk", linked_text))
        setattr(self, "popup_c3_linked", linked_text)

        if CollapsibleSection:
            linked_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            linked_section.set_content_height(140)
            linked_section._min_height = 80
            linked_section._max_height = 220
            linked_section.set_header_style("""
                QFrame {
                    background: rgba(22, 163, 74, 0.15);
                    border: 1px solid rgba(22, 163, 74, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            linked_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #15803d;
                    background: transparent;
                    border: none;
                }
            """)

            linked_content = QWidget()
            linked_content.setStyleSheet("""
                QWidget {
                    background: rgba(240, 253, 244, 0.95);
                    border: 1px solid rgba(22, 163, 74, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            linked_cb_layout = QVBoxLayout(linked_content)
            linked_cb_layout.setContentsMargins(10, 6, 10, 6)
            linked_cb_layout.setSpacing(3)

            C3_LINKED_ITEMS = [
                ("lnk_violence_psychosis", "Violence occurred during psychosis"),
                ("lnk_violence_relapse", "Violence linked to relapse"),
                ("lnk_aggression_unwell", "Aggression increases when unwell"),
                ("lnk_remission", "Symptoms currently in remission (protective)"),
                ("lnk_stable_medication", "Stable on medication (protective)"),
            ]

            c3_linked_checkboxes = []
            for item_key, item_label in C3_LINKED_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c3_key", item_key)
                cb.setProperty("c3_category", "linked")
                cb.stateChanged.connect(self._update_c3_symptoms_narrative)
                linked_cb_layout.addWidget(cb)
                c3_linked_checkboxes.append(cb)

            linked_section.set_content(linked_content)
            linked_layout.addWidget(linked_section)
            setattr(self, "popup_c3_linked_checkboxes", c3_linked_checkboxes)

        subsections_target_layout.addWidget(linked_container)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")

            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_c3_symptoms_narrative(self):
        """Update C3 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. PSYCHOTIC SYMPTOMS ============
        psy_items = {
            "psy_paranoid": "paranoid ideation",
            "psy_persecutory": "persecutory delusions",
            "psy_command_hallucinations": "command hallucinations",
            "psy_hearing_voices": "hearing voices",
            "psy_grandiose": "grandiose delusions",
            "psy_thought_disorder": "thought disorder",
            "psy_actively_psychotic": "being actively psychotic",
        }

        psy_cbs = getattr(self, "popup_c3_psychotic_checkboxes", [])
        psy_text = getattr(self, "popup_c3_psychotic", None)

        psy_checked = []
        for cb in psy_cbs:
            if cb.isChecked():
                key = cb.property("c3_key")
                if key in psy_items:
                    psy_checked.append(psy_items[key])

        if psy_text:
            if psy_checked:
                psy_text.setPlainText(f"{Subj} is currently experiencing psychotic symptoms, including {join_items(psy_checked)}.")
            else:
                psy_text.clear()

        # ============ 2. MANIA/HYPOMANIA ============
        man_items = {
            "man_manic": "a manic episode",
            "man_hypomanic": "a hypomanic episode",
            "man_elevated_mood": "elevated mood",
            "man_grandiosity": "grandiosity",
            "man_disinhibited": "disinhibition",
            "man_reduced_sleep": "reduced need for sleep",
        }

        man_cbs = getattr(self, "popup_c3_mania_checkboxes", [])
        man_text = getattr(self, "popup_c3_mania", None)

        man_checked = []
        for cb in man_cbs:
            if cb.isChecked():
                key = cb.property("c3_key")
                if key in man_items:
                    man_checked.append(man_items[key])

        if man_text:
            if man_checked:
                man_text.setPlainText(f"There are features of mania or hypomania present, including {join_items(man_checked)}.")
            else:
                man_text.clear()

        # ============ 3. SEVERE DEPRESSION ============
        dep_items = {
            "dep_severe": "severe depression",
            "dep_agitated": "agitated depression",
            "dep_hopelessness": "hopelessness with anger",
            "dep_nihilistic": "nihilistic beliefs",
            "dep_paranoid": "paranoid depression",
        }

        dep_cbs = getattr(self, "popup_c3_depression_checkboxes", [])
        dep_text = getattr(self, "popup_c3_depression", None)

        dep_checked = []
        for cb in dep_cbs:
            if cb.isChecked():
                key = cb.property("c3_key")
                if key in dep_items:
                    dep_checked.append(dep_items[key])

        if dep_text:
            if dep_checked:
                dep_text.setPlainText(f"{Subj} is presenting with depressive symptoms relevant to risk, including {join_items(dep_checked)}.")
            else:
                dep_text.clear()

        # ============ 4. AFFECTIVE INSTABILITY ============
        aff_items = {
            "aff_labile": "labile affect",
            "aff_easily_provoked": "being easily provoked",
            "aff_low_frustration": "low frustration tolerance",
            "aff_explosive": "explosive anger",
            "aff_rapid_shifts": "rapid mood shifts",
        }

        aff_cbs = getattr(self, "popup_c3_affective_checkboxes", [])
        aff_text = getattr(self, "popup_c3_affective", None)

        aff_checked = []
        for cb in aff_cbs:
            if cb.isChecked():
                key = cb.property("c3_key")
                if key in aff_items:
                    aff_checked.append(aff_items[key])

        if aff_text:
            if aff_checked:
                aff_text.setPlainText(f"{Subj} demonstrates affective instability, with {join_items(aff_checked)}.")
            else:
                aff_text.clear()

        # ============ 5. AROUSAL/ANXIETY STATES ============
        ars_items = {
            "ars_hypervigilant": "hypervigilance",
            "ars_on_edge": "being on edge and tense",
            "ars_threat_perception": "heightened threat perception",
            "ars_ptsd_exacerbated": "exacerbated PTSD symptoms",
        }

        ars_cbs = getattr(self, "popup_c3_arousal_checkboxes", [])
        ars_text = getattr(self, "popup_c3_arousal", None)

        ars_checked = []
        for cb in ars_cbs:
            if cb.isChecked():
                key = cb.property("c3_key")
                if key in ars_items:
                    ars_checked.append(ars_items[key])

        if ars_text:
            if ars_checked:
                ars_text.setPlainText(f"There is evidence of heightened arousal, including {join_items(ars_checked)}.")
            else:
                ars_text.clear()

        # ============ 6. SYMPTOMS LINKED TO VIOLENCE (with protective factors) ============
        lnk_vuln_items = {
            "lnk_violence_psychosis": "violence during psychosis",
            "lnk_violence_relapse": "violence linked to relapse",
            "lnk_aggression_unwell": "increased aggression when unwell",
        }
        lnk_prot_items = {
            "lnk_remission": "symptoms currently in remission",
            "lnk_stable_medication": "being stable on medication",
        }

        lnk_cbs = getattr(self, "popup_c3_linked_checkboxes", [])
        lnk_text = getattr(self, "popup_c3_linked", None)

        lnk_vuln_checked = []
        lnk_prot_checked = []
        for cb in lnk_cbs:
            if cb.isChecked():
                key = cb.property("c3_key")
                if key in lnk_vuln_items:
                    lnk_vuln_checked.append(lnk_vuln_items[key])
                elif key in lnk_prot_items:
                    lnk_prot_checked.append(lnk_prot_items[key])

        if lnk_text:
            sentences = []
            if lnk_vuln_checked:
                sentences.append(f"There is a documented link between symptoms and violence, with a history of {join_items(lnk_vuln_checked)}.")
            if lnk_prot_checked:
                if sentences:
                    sentences.append(f"However, protective factors include {join_items(lnk_prot_checked)}.")
                else:
                    sentences.append(f"Protective factors include {join_items(lnk_prot_checked)}.")
            if sentences:
                lnk_text.setPlainText(" ".join(sentences))
            else:
                lnk_text.clear()

    def _build_hcr_c4_popup(self, key: str, code: str, description: str, subsections: list):
        """Build C4 (Instability) popup with tick boxes.

        Based on HCR-20V3 manual guidance for C4 - affective instability,
        behavioural instability, and cognitive instability.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. AFFECTIVE INSTABILITY CONTAINER ===
        affective_container = QFrame()
        affective_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        affective_layout = QVBoxLayout(affective_container)
        affective_layout.setContentsMargins(12, 10, 12, 10)
        affective_layout.setSpacing(8)

        affective_lbl = QLabel("Affective Instability")
        affective_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        affective_layout.addWidget(affective_lbl)

        affective_text = QTextEdit()
        affective_text.setPlaceholderText("Describe affective instability (mood swings, emotional dysregulation)...")
        affective_text.setStyleSheet(input_style + " background: white;")
        affective_text.setMinimumHeight(58)
        affective_layout.addWidget(affective_text)
        affective_layout.addWidget(DragResizeBar(affective_text))
        text_widgets.append(affective_text)
        subsection_data.append(("Affective Instability", affective_text))
        setattr(self, "popup_c4_affective", affective_text)

        if CollapsibleSection:
            affective_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            affective_section.set_content_height(160)
            affective_section._min_height = 80
            affective_section._max_height = 250
            affective_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            affective_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            affective_content = QWidget()
            affective_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            affective_cb_layout = QVBoxLayout(affective_content)
            affective_cb_layout.setContentsMargins(10, 6, 10, 6)
            affective_cb_layout.setSpacing(3)

            C4_AFFECTIVE_ITEMS = [
                ("aff_mood_swings", "Mood swings"),
                ("aff_volatile", "Volatile mood"),
                ("aff_labile", "Labile affect"),
                ("aff_irritable", "Irritable"),
                ("aff_easily_angered", "Easily angered"),
                ("aff_emotional_dysreg", "Poor emotional regulation"),
            ]

            c4_affective_checkboxes = []
            for item_key, item_label in C4_AFFECTIVE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c4_key", item_key)
                cb.setProperty("c4_category", "affective")
                cb.stateChanged.connect(self._update_c4_instability_narrative)
                affective_cb_layout.addWidget(cb)
                c4_affective_checkboxes.append(cb)

            affective_section.set_content(affective_content)
            affective_layout.addWidget(affective_section)
            setattr(self, "popup_c4_affective_checkboxes", c4_affective_checkboxes)

        subsections_target_layout.addWidget(affective_container)

        # === 2. BEHAVIOURAL IMPULSIVITY CONTAINER ===
        impulsive_container = QFrame()
        impulsive_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        impulsive_layout = QVBoxLayout(impulsive_container)
        impulsive_layout.setContentsMargins(12, 10, 12, 10)
        impulsive_layout.setSpacing(8)

        impulsive_lbl = QLabel("Behavioural Impulsivity")
        impulsive_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #b45309; background: transparent; border: none;")
        impulsive_layout.addWidget(impulsive_lbl)

        impulsive_text = QTextEdit()
        impulsive_text.setPlaceholderText("Describe behavioural impulsivity (acts without thinking, poor impulse control)...")
        impulsive_text.setStyleSheet(input_style + " background: white;")
        impulsive_text.setMinimumHeight(58)
        impulsive_layout.addWidget(impulsive_text)
        impulsive_layout.addWidget(DragResizeBar(impulsive_text))
        text_widgets.append(impulsive_text)
        subsection_data.append(("Behavioural Impulsivity", impulsive_text))
        setattr(self, "popup_c4_impulsive", impulsive_text)

        if CollapsibleSection:
            impulsive_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            impulsive_section.set_content_height(160)
            impulsive_section._min_height = 80
            impulsive_section._max_height = 250
            impulsive_section.set_header_style("""
                QFrame {
                    background: rgba(217, 119, 6, 0.15);
                    border: 1px solid rgba(217, 119, 6, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            impulsive_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #b45309;
                    background: transparent;
                    border: none;
                }
            """)

            impulsive_content = QWidget()
            impulsive_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(217, 119, 6, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            impulsive_cb_layout = QVBoxLayout(impulsive_content)
            impulsive_cb_layout.setContentsMargins(10, 6, 10, 6)
            impulsive_cb_layout.setSpacing(3)

            C4_IMPULSIVE_ITEMS = [
                ("imp_acts_without_thinking", "Acts without thinking"),
                ("imp_poor_impulse_control", "Poor impulse control"),
                ("imp_unpredictable", "Unpredictable behaviour"),
                ("imp_erratic", "Erratic behaviour"),
                ("imp_reckless", "Reckless behaviour"),
            ]

            c4_impulsive_checkboxes = []
            for item_key, item_label in C4_IMPULSIVE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c4_key", item_key)
                cb.setProperty("c4_category", "impulsive")
                cb.stateChanged.connect(self._update_c4_instability_narrative)
                impulsive_cb_layout.addWidget(cb)
                c4_impulsive_checkboxes.append(cb)

            impulsive_section.set_content(impulsive_content)
            impulsive_layout.addWidget(impulsive_section)
            setattr(self, "popup_c4_impulsive_checkboxes", c4_impulsive_checkboxes)

        subsections_target_layout.addWidget(impulsive_container)

        # === 3. ANGER MANAGEMENT CONTAINER ===
        anger_container = QFrame()
        anger_container.setStyleSheet("""
            QFrame {
                background: rgba(243, 232, 255, 0.5);
                border: 2px solid rgba(147, 51, 234, 0.4);
                border-radius: 10px;
            }
        """)
        anger_layout = QVBoxLayout(anger_container)
        anger_layout.setContentsMargins(12, 10, 12, 10)
        anger_layout.setSpacing(8)

        anger_lbl = QLabel("Anger Dyscontrol")
        anger_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #7e22ce; background: transparent; border: none;")
        anger_layout.addWidget(anger_lbl)

        anger_text = QTextEdit()
        anger_text.setPlaceholderText("Describe anger management difficulties (explosive outbursts, difficulty controlling temper)...")
        anger_text.setStyleSheet(input_style + " background: white;")
        anger_text.setMinimumHeight(58)
        anger_layout.addWidget(anger_text)
        anger_layout.addWidget(DragResizeBar(anger_text))
        text_widgets.append(anger_text)
        subsection_data.append(("Anger Dyscontrol", anger_text))
        setattr(self, "popup_c4_anger", anger_text)

        if CollapsibleSection:
            anger_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            anger_section.set_content_height(140)
            anger_section._min_height = 80
            anger_section._max_height = 220
            anger_section.set_header_style("""
                QFrame {
                    background: rgba(147, 51, 234, 0.15);
                    border: 1px solid rgba(147, 51, 234, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            anger_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #7e22ce;
                    background: transparent;
                    border: none;
                }
            """)

            anger_content = QWidget()
            anger_content.setStyleSheet("""
                QWidget {
                    background: rgba(250, 245, 255, 0.95);
                    border: 1px solid rgba(147, 51, 234, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            anger_cb_layout = QVBoxLayout(anger_content)
            anger_cb_layout.setContentsMargins(10, 6, 10, 6)
            anger_cb_layout.setSpacing(3)

            C4_ANGER_ITEMS = [
                ("ang_explosive", "Explosive outbursts"),
                ("ang_angry_outburst", "Angry outbursts"),
                ("ang_difficulty_temper", "Difficulty controlling temper"),
                ("ang_agitated", "Frequently agitated"),
                ("ang_hostile", "Hostile manner"),
            ]

            c4_anger_checkboxes = []
            for item_key, item_label in C4_ANGER_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c4_key", item_key)
                cb.setProperty("c4_category", "anger")
                cb.stateChanged.connect(self._update_c4_instability_narrative)
                anger_cb_layout.addWidget(cb)
                c4_anger_checkboxes.append(cb)

            anger_section.set_content(anger_content)
            anger_layout.addWidget(anger_section)
            setattr(self, "popup_c4_anger_checkboxes", c4_anger_checkboxes)

        subsections_target_layout.addWidget(anger_container)

        # === 4. ENVIRONMENTAL INSTABILITY CONTAINER ===
        environ_container = QFrame()
        environ_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        environ_layout = QVBoxLayout(environ_container)
        environ_layout.setContentsMargins(12, 10, 12, 10)
        environ_layout.setSpacing(8)

        environ_lbl = QLabel("Environmental/Life Instability")
        environ_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        environ_layout.addWidget(environ_lbl)

        environ_text = QTextEdit()
        environ_text.setPlaceholderText("Describe recent life instability (relationship changes, housing, employment)...")
        environ_text.setStyleSheet(input_style + " background: white;")
        environ_text.setMinimumHeight(58)
        environ_layout.addWidget(environ_text)
        environ_layout.addWidget(DragResizeBar(environ_text))
        text_widgets.append(environ_text)
        subsection_data.append(("Environmental/Life Instability", environ_text))
        setattr(self, "popup_c4_environ", environ_text)

        if CollapsibleSection:
            environ_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            environ_section.set_content_height(140)
            environ_section._min_height = 80
            environ_section._max_height = 220
            environ_section.set_header_style("""
                QFrame {
                    background: rgba(59, 130, 246, 0.15);
                    border: 1px solid rgba(59, 130, 246, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            environ_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #1d4ed8;
                    background: transparent;
                    border: none;
                }
            """)

            environ_content = QWidget()
            environ_content.setStyleSheet("""
                QWidget {
                    background: rgba(239, 246, 255, 0.95);
                    border: 1px solid rgba(59, 130, 246, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            environ_cb_layout = QVBoxLayout(environ_content)
            environ_cb_layout.setContentsMargins(10, 6, 10, 6)
            environ_cb_layout.setSpacing(3)

            C4_ENVIRON_ITEMS = [
                ("env_relationship_breakdown", "Relationship breakdown"),
                ("env_housing_instability", "Housing instability"),
                ("env_job_loss", "Job loss / unemployment"),
                ("env_financial_crisis", "Financial crisis"),
                ("env_recent_move", "Recent move / relocation"),
            ]

            c4_environ_checkboxes = []
            for item_key, item_label in C4_ENVIRON_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c4_key", item_key)
                cb.setProperty("c4_category", "environ")
                cb.stateChanged.connect(self._update_c4_instability_narrative)
                environ_cb_layout.addWidget(cb)
                c4_environ_checkboxes.append(cb)

            environ_section.set_content(environ_content)
            environ_layout.addWidget(environ_section)
            setattr(self, "popup_c4_environ_checkboxes", c4_environ_checkboxes)

        subsections_target_layout.addWidget(environ_container)

        # === 5. STABILITY INDICATORS (PROTECTIVE) CONTAINER ===
        stability_container = QFrame()
        stability_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(22, 163, 74, 0.4);
                border-radius: 10px;
            }
        """)
        stability_layout = QVBoxLayout(stability_container)
        stability_layout.setContentsMargins(12, 10, 12, 10)
        stability_layout.setSpacing(8)

        stability_lbl = QLabel("Stability Indicators (Protective)")
        stability_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        stability_layout.addWidget(stability_lbl)

        stability_text = QTextEdit()
        stability_text.setPlaceholderText("Describe evidence of stability (settled lifestyle, good emotional regulation)...")
        stability_text.setStyleSheet(input_style + " background: white;")
        stability_text.setMinimumHeight(58)
        stability_layout.addWidget(stability_text)
        stability_layout.addWidget(DragResizeBar(stability_text))
        text_widgets.append(stability_text)
        subsection_data.append(("Stability Indicators", stability_text))
        setattr(self, "popup_c4_stability", stability_text)

        if CollapsibleSection:
            stability_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            stability_section.set_content_height(120)
            stability_section._min_height = 80
            stability_section._max_height = 200
            stability_section.set_header_style("""
                QFrame {
                    background: rgba(22, 163, 74, 0.15);
                    border: 1px solid rgba(22, 163, 74, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            stability_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #15803d;
                    background: transparent;
                    border: none;
                }
            """)

            stability_content = QWidget()
            stability_content.setStyleSheet("""
                QWidget {
                    background: rgba(240, 253, 244, 0.95);
                    border: 1px solid rgba(22, 163, 74, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            stability_cb_layout = QVBoxLayout(stability_content)
            stability_cb_layout.setContentsMargins(10, 6, 10, 6)
            stability_cb_layout.setSpacing(3)

            C4_STABILITY_ITEMS = [
                ("stab_good_emotional_reg", "Good emotional regulation"),
                ("stab_stable_mood", "Stable mood"),
                ("stab_settled_lifestyle", "Settled lifestyle"),
                ("stab_consistent_routine", "Consistent daily routine"),
            ]

            c4_stability_checkboxes = []
            for item_key, item_label in C4_STABILITY_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c4_key", item_key)
                cb.setProperty("c4_category", "stability")
                cb.stateChanged.connect(self._update_c4_instability_narrative)
                stability_cb_layout.addWidget(cb)
                c4_stability_checkboxes.append(cb)

            stability_section.set_content(stability_content)
            stability_layout.addWidget(stability_section)
            setattr(self, "popup_c4_stability_checkboxes", c4_stability_checkboxes)

        subsections_target_layout.addWidget(stability_container)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")

            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_c4_instability_narrative(self):
        """Update C4 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. AFFECTIVE INSTABILITY ============
        aff_items = {
            "aff_mood_swings": "mood swings",
            "aff_volatile": "volatile mood",
            "aff_labile": "labile affect",
            "aff_irritable": "irritability",
            "aff_easily_angered": "being easily angered",
            "aff_emotional_dysreg": "poor emotional regulation",
        }

        aff_cbs = getattr(self, "popup_c4_affective_checkboxes", [])
        aff_text = getattr(self, "popup_c4_affective", None)

        aff_checked = []
        for cb in aff_cbs:
            if cb.isChecked():
                key = cb.property("c4_key")
                if key in aff_items:
                    aff_checked.append(aff_items[key])

        if aff_text:
            if aff_checked:
                aff_text.setPlainText(f"{Subj} demonstrates affective instability, with {join_items(aff_checked)}.")
            else:
                aff_text.clear()

        # ============ 2. BEHAVIOURAL IMPULSIVITY ============
        imp_items = {
            "imp_acts_without_thinking": "acting without thinking",
            "imp_poor_impulse_control": "poor impulse control",
            "imp_unpredictable": "unpredictable behaviour",
            "imp_erratic": "erratic behaviour",
            "imp_reckless": "reckless behaviour",
        }

        imp_cbs = getattr(self, "popup_c4_impulsive_checkboxes", [])
        imp_text = getattr(self, "popup_c4_impulsive", None)

        imp_checked = []
        for cb in imp_cbs:
            if cb.isChecked():
                key = cb.property("c4_key")
                if key in imp_items:
                    imp_checked.append(imp_items[key])

        if imp_text:
            if imp_checked:
                imp_text.setPlainText(f"There is evidence of behavioural impulsivity, including {join_items(imp_checked)}.")
            else:
                imp_text.clear()

        # ============ 3. ANGER DYSCONTROL ============
        ang_items = {
            "ang_explosive": "explosive outbursts",
            "ang_angry_outburst": "angry outbursts",
            "ang_difficulty_temper": "difficulty controlling {poss} temper",
            "ang_agitated": "frequent agitation",
            "ang_hostile": "a hostile manner",
        }

        ang_cbs = getattr(self, "popup_c4_anger_checkboxes", [])
        ang_text = getattr(self, "popup_c4_anger", None)

        ang_checked = []
        for cb in ang_cbs:
            if cb.isChecked():
                key = cb.property("c4_key")
                if key in ang_items:
                    ang_checked.append(ang_items[key].format(poss=poss))

        if ang_text:
            if ang_checked:
                ang_text.setPlainText(f"{Subj} has difficulty managing anger, with {join_items(ang_checked)}.")
            else:
                ang_text.clear()

        # ============ 4. ENVIRONMENTAL/LIFE INSTABILITY ============
        env_items = {
            "env_relationship_breakdown": "relationship breakdown",
            "env_housing_instability": "housing instability",
            "env_job_loss": "job loss or unemployment",
            "env_financial_crisis": "financial crisis",
            "env_recent_move": "recent relocation",
        }

        env_cbs = getattr(self, "popup_c4_environ_checkboxes", [])
        env_text = getattr(self, "popup_c4_environ", None)

        env_checked = []
        for cb in env_cbs:
            if cb.isChecked():
                key = cb.property("c4_key")
                if key in env_items:
                    env_checked.append(env_items[key])

        if env_text:
            if env_checked:
                env_text.setPlainText(f"Recent life events have contributed to instability, including {join_items(env_checked)}.")
            else:
                env_text.clear()

        # ============ 5. STABILITY INDICATORS (PROTECTIVE) ============
        stab_items = {
            "stab_good_emotional_reg": "good emotional regulation",
            "stab_stable_mood": "stable mood",
            "stab_settled_lifestyle": "a settled lifestyle",
            "stab_consistent_routine": "a consistent daily routine",
        }

        stab_cbs = getattr(self, "popup_c4_stability_checkboxes", [])
        stab_text = getattr(self, "popup_c4_stability", None)

        stab_checked = []
        for cb in stab_cbs:
            if cb.isChecked():
                key = cb.property("c4_key")
                if key in stab_items:
                    stab_checked.append(stab_items[key])

        if stab_text:
            if stab_checked:
                stab_text.setPlainText(f"Protective factors include {join_items(stab_checked)}.")
            else:
                stab_text.clear()

    def _build_hcr_c5_popup(self, key: str, code: str, description: str, subsections: list):
        """Build C5 (Treatment or Supervision Response) popup with tick boxes.

        Based on HCR-20V3 manual guidance for C5 - response to treatment,
        medication adherence, engagement with services, compliance with conditions.
        """
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        # Item description
        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        # Relevance scoring
        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. MEDICATION ADHERENCE CONTAINER ===
        medication_container = QFrame()
        medication_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        medication_layout = QVBoxLayout(medication_container)
        medication_layout.setContentsMargins(12, 10, 12, 10)
        medication_layout.setSpacing(8)

        medication_lbl = QLabel("Medication Adherence")
        medication_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        medication_layout.addWidget(medication_lbl)

        medication_text = QTextEdit()
        medication_text.setPlaceholderText("Describe medication adherence, concordance, response to treatment...")
        medication_text.setStyleSheet(input_style + " background: white;")
        medication_text.setMinimumHeight(58)
        medication_layout.addWidget(medication_text)
        medication_layout.addWidget(DragResizeBar(medication_text))
        text_widgets.append(medication_text)
        subsection_data.append(("Medication Adherence", medication_text))
        setattr(self, "popup_c5_medication", medication_text)

        if CollapsibleSection:
            medication_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            medication_section.set_content_height(180)
            medication_section._min_height = 80
            medication_section._max_height = 280
            medication_section.set_header_style("""
                QFrame {
                    background: rgba(220, 38, 38, 0.15);
                    border: 1px solid rgba(220, 38, 38, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            medication_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #991b1b;
                    background: transparent;
                    border: none;
                }
            """)

            medication_content = QWidget()
            medication_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 242, 242, 0.95);
                    border: 1px solid rgba(220, 38, 38, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            medication_cb_layout = QVBoxLayout(medication_content)
            medication_cb_layout.setContentsMargins(10, 6, 10, 6)
            medication_cb_layout.setSpacing(3)

            C5_MEDICATION_ITEMS = [
                ("med_non_compliant", "Non-compliant with medication"),
                ("med_stops_discharge", "Stops medication after discharge"),
                ("med_refuses", "Refuses medication"),
                ("med_selective", "Selective/partial adherence"),
                ("med_covert_non_compliance", "Covert non-compliance"),
                ("med_accepts", "Accepts medication (protective)"),
                ("med_consistent", "Consistently takes medication (protective)"),
            ]

            c5_medication_checkboxes = []
            for item_key, item_label in C5_MEDICATION_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c5_key", item_key)
                cb.setProperty("c5_category", "medication")
                cb.stateChanged.connect(self._update_c5_treatment_narrative)
                medication_cb_layout.addWidget(cb)
                c5_medication_checkboxes.append(cb)

            medication_section.set_content(medication_content)
            medication_layout.addWidget(medication_section)
            setattr(self, "popup_c5_medication_checkboxes", c5_medication_checkboxes)

        subsections_target_layout.addWidget(medication_container)

        # === 2. ENGAGEMENT WITH SERVICES CONTAINER ===
        engagement_container = QFrame()
        engagement_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        engagement_layout = QVBoxLayout(engagement_container)
        engagement_layout.setContentsMargins(12, 10, 12, 10)
        engagement_layout.setSpacing(8)

        engagement_lbl = QLabel("Engagement with Services")
        engagement_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #b45309; background: transparent; border: none;")
        engagement_layout.addWidget(engagement_lbl)

        engagement_text = QTextEdit()
        engagement_text.setPlaceholderText("Describe engagement with services, attendance, participation in care...")
        engagement_text.setStyleSheet(input_style + " background: white;")
        engagement_text.setMinimumHeight(58)
        engagement_layout.addWidget(engagement_text)
        engagement_layout.addWidget(DragResizeBar(engagement_text))
        text_widgets.append(engagement_text)
        subsection_data.append(("Engagement with Services", engagement_text))
        setattr(self, "popup_c5_engagement", engagement_text)

        if CollapsibleSection:
            engagement_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            engagement_section.set_content_height(160)
            engagement_section._min_height = 80
            engagement_section._max_height = 250
            engagement_section.set_header_style("""
                QFrame {
                    background: rgba(217, 119, 6, 0.15);
                    border: 1px solid rgba(217, 119, 6, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            engagement_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #b45309;
                    background: transparent;
                    border: none;
                }
            """)

            engagement_content = QWidget()
            engagement_content.setStyleSheet("""
                QWidget {
                    background: rgba(254, 252, 232, 0.95);
                    border: 1px solid rgba(217, 119, 6, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            engagement_cb_layout = QVBoxLayout(engagement_content)
            engagement_cb_layout.setContentsMargins(10, 6, 10, 6)
            engagement_cb_layout.setSpacing(3)

            C5_ENGAGEMENT_ITEMS = [
                ("eng_disengaged", "Disengaged from services"),
                ("eng_misses_appts", "Misses appointments (DNA)"),
                ("eng_poor_attendance", "Poor attendance"),
                ("eng_avoids_reviews", "Avoids reviews"),
                ("eng_actively_engages", "Actively engages with services (protective)"),
            ]

            c5_engagement_checkboxes = []
            for item_key, item_label in C5_ENGAGEMENT_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c5_key", item_key)
                cb.setProperty("c5_category", "engagement")
                cb.stateChanged.connect(self._update_c5_treatment_narrative)
                engagement_cb_layout.addWidget(cb)
                c5_engagement_checkboxes.append(cb)

            engagement_section.set_content(engagement_content)
            engagement_layout.addWidget(engagement_section)
            setattr(self, "popup_c5_engagement_checkboxes", c5_engagement_checkboxes)

        subsections_target_layout.addWidget(engagement_container)

        # === 3. COMPLIANCE WITH CONDITIONS CONTAINER ===
        compliance_container = QFrame()
        compliance_container.setStyleSheet("""
            QFrame {
                background: rgba(243, 232, 255, 0.5);
                border: 2px solid rgba(147, 51, 234, 0.4);
                border-radius: 10px;
            }
        """)
        compliance_layout = QVBoxLayout(compliance_container)
        compliance_layout.setContentsMargins(12, 10, 12, 10)
        compliance_layout.setSpacing(8)

        compliance_lbl = QLabel("Compliance with Conditions")
        compliance_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #7e22ce; background: transparent; border: none;")
        compliance_layout.addWidget(compliance_lbl)

        compliance_text = QTextEdit()
        compliance_text.setPlaceholderText("Describe compliance with CTO, conditional discharge, probation, licence conditions...")
        compliance_text.setStyleSheet(input_style + " background: white;")
        compliance_text.setMinimumHeight(58)
        compliance_layout.addWidget(compliance_text)
        compliance_layout.addWidget(DragResizeBar(compliance_text))
        text_widgets.append(compliance_text)
        subsection_data.append(("Compliance with Conditions", compliance_text))
        setattr(self, "popup_c5_compliance", compliance_text)

        if CollapsibleSection:
            compliance_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            compliance_section.set_content_height(160)
            compliance_section._min_height = 80
            compliance_section._max_height = 250
            compliance_section.set_header_style("""
                QFrame {
                    background: rgba(147, 51, 234, 0.15);
                    border: 1px solid rgba(147, 51, 234, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            compliance_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #7e22ce;
                    background: transparent;
                    border: none;
                }
            """)

            compliance_content = QWidget()
            compliance_content.setStyleSheet("""
                QWidget {
                    background: rgba(250, 245, 255, 0.95);
                    border: 1px solid rgba(147, 51, 234, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            compliance_cb_layout = QVBoxLayout(compliance_content)
            compliance_cb_layout.setContentsMargins(10, 6, 10, 6)
            compliance_cb_layout.setSpacing(3)

            C5_COMPLIANCE_ITEMS = [
                ("cmp_breaches", "Breaches conditions"),
                ("cmp_absconded", "Has absconded"),
                ("cmp_recalled", "Required recall"),
                ("cmp_only_coerced", "Only complies under coercion"),
                ("cmp_resists_monitoring", "Resists monitoring"),
                ("cmp_accepts_conditions", "Accepts conditions (protective)"),
            ]

            c5_compliance_checkboxes = []
            for item_key, item_label in C5_COMPLIANCE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c5_key", item_key)
                cb.setProperty("c5_category", "compliance")
                cb.stateChanged.connect(self._update_c5_treatment_narrative)
                compliance_cb_layout.addWidget(cb)
                c5_compliance_checkboxes.append(cb)

            compliance_section.set_content(compliance_content)
            compliance_layout.addWidget(compliance_section)
            setattr(self, "popup_c5_compliance_checkboxes", c5_compliance_checkboxes)

        subsections_target_layout.addWidget(compliance_container)

        # === 4. PATTERN OVER TIME CONTAINER ===
        pattern_container = QFrame()
        pattern_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        pattern_layout = QVBoxLayout(pattern_container)
        pattern_layout.setContentsMargins(12, 10, 12, 10)
        pattern_layout.setSpacing(8)

        pattern_lbl = QLabel("Pattern Over Time")
        pattern_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        pattern_layout.addWidget(pattern_lbl)

        pattern_text = QTextEdit()
        pattern_text.setPlaceholderText("Describe patterns of compliance/non-compliance over time...")
        pattern_text.setStyleSheet(input_style + " background: white;")
        pattern_text.setMinimumHeight(58)
        pattern_layout.addWidget(pattern_text)
        pattern_layout.addWidget(DragResizeBar(pattern_text))
        text_widgets.append(pattern_text)
        subsection_data.append(("Pattern Over Time", pattern_text))
        setattr(self, "popup_c5_pattern", pattern_text)

        if CollapsibleSection:
            pattern_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            pattern_section.set_content_height(140)
            pattern_section._min_height = 80
            pattern_section._max_height = 220
            pattern_section.set_header_style("""
                QFrame {
                    background: rgba(59, 130, 246, 0.15);
                    border: 1px solid rgba(59, 130, 246, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            pattern_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #1d4ed8;
                    background: transparent;
                    border: none;
                }
            """)

            pattern_content = QWidget()
            pattern_content.setStyleSheet("""
                QWidget {
                    background: rgba(239, 246, 255, 0.95);
                    border: 1px solid rgba(59, 130, 246, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            pattern_cb_layout = QVBoxLayout(pattern_content)
            pattern_cb_layout.setContentsMargins(10, 6, 10, 6)
            pattern_cb_layout.setSpacing(3)

            C5_PATTERN_ITEMS = [
                ("pat_repeated_disengage", "Repeated disengagement"),
                ("pat_history_non_compliance", "History of non-compliance"),
                ("pat_cycle_relapse", "Cycle of engagement → discharge → disengagement → relapse"),
                ("pat_sustained_adherence", "Sustained adherence over time (protective)"),
            ]

            c5_pattern_checkboxes = []
            for item_key, item_label in C5_PATTERN_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c5_key", item_key)
                cb.setProperty("c5_category", "pattern")
                cb.stateChanged.connect(self._update_c5_treatment_narrative)
                pattern_cb_layout.addWidget(cb)
                c5_pattern_checkboxes.append(cb)

            pattern_section.set_content(pattern_content)
            pattern_layout.addWidget(pattern_section)
            setattr(self, "popup_c5_pattern_checkboxes", c5_pattern_checkboxes)

        subsections_target_layout.addWidget(pattern_container)

        # === 5. TREATMENT RESPONSIVENESS CONTAINER ===
        response_container = QFrame()
        response_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(22, 163, 74, 0.4);
                border-radius: 10px;
            }
        """)
        response_layout = QVBoxLayout(response_container)
        response_layout.setContentsMargins(12, 10, 12, 10)
        response_layout.setSpacing(8)

        response_lbl = QLabel("Treatment Responsiveness")
        response_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        response_layout.addWidget(response_lbl)

        response_text = QTextEdit()
        response_text.setPlaceholderText("Describe response to treatment, effectiveness of interventions...")
        response_text.setStyleSheet(input_style + " background: white;")
        response_text.setMinimumHeight(58)
        response_layout.addWidget(response_text)
        response_layout.addWidget(DragResizeBar(response_text))
        text_widgets.append(response_text)
        subsection_data.append(("Treatment Responsiveness", response_text))
        setattr(self, "popup_c5_response", response_text)

        if CollapsibleSection:
            response_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            response_section.set_content_height(120)
            response_section._min_height = 80
            response_section._max_height = 200
            response_section.set_header_style("""
                QFrame {
                    background: rgba(22, 163, 74, 0.15);
                    border: 1px solid rgba(22, 163, 74, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            response_section.set_title_style("""
                QLabel {
                    font-size: 14px;
                    font-weight: 600;
                    color: #15803d;
                    background: transparent;
                    border: none;
                }
            """)

            response_content = QWidget()
            response_content.setStyleSheet("""
                QWidget {
                    background: rgba(240, 253, 244, 0.95);
                    border: 1px solid rgba(22, 163, 74, 0.2);
                    border-top: none;
                    border-radius: 0 0 8px 8px;
                }
                QCheckBox {
                    background: transparent;
                    border: none;
                    font-size: 13px;
                    color: #374151;
                    padding: 2px;
                }
            """)

            response_cb_layout = QVBoxLayout(response_content)
            response_cb_layout.setContentsMargins(10, 6, 10, 6)
            response_cb_layout.setSpacing(3)

            C5_RESPONSE_ITEMS = [
                ("rsp_treatment_resistant", "Treatment resistant"),
                ("rsp_no_improvement", "No improvement with treatment"),
                ("rsp_responds_well", "Responds well to treatment (protective)"),
                ("rsp_benefits_therapy", "Benefits from psychological therapy (protective)"),
            ]

            c5_response_checkboxes = []
            for item_key, item_label in C5_RESPONSE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("c5_key", item_key)
                cb.setProperty("c5_category", "response")
                cb.stateChanged.connect(self._update_c5_treatment_narrative)
                response_cb_layout.addWidget(cb)
                c5_response_checkboxes.append(cb)

            response_section.set_content(response_content)
            response_layout.addWidget(response_section)
            setattr(self, "popup_c5_response_checkboxes", c5_response_checkboxes)

        subsections_target_layout.addWidget(response_container)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400
            imported_section.set_header_style("""
                QFrame {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 6px 6px 0 0;
                }
            """)
            imported_section.set_title_style("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            imported_content = QWidget()
            imported_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 10px 10px;
                }
            """)

            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)
            imported_content_layout.setSpacing(4)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]

            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")

            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")

            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_c5_treatment_narrative(self):
        """Update C5 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss, refl = 'he', 'him', 'his', 'himself'
        else:
            subj, obj, poss, refl = 'she', 'her', 'her', 'herself'
        Subj = subj.capitalize()
        Poss = poss.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. MEDICATION ADHERENCE ============
        med_vuln_items = {
            "med_non_compliant": "non-compliance",
            "med_stops_discharge": "stopping medication after discharge",
            "med_refuses": "refusing medication",
            "med_selective": "selective or partial adherence",
            "med_covert_non_compliance": "covert non-compliance",
        }
        med_prot_items = {
            "med_accepts": "accepting medication",
            "med_consistent": "consistent adherence",
        }

        med_cbs = getattr(self, "popup_c5_medication_checkboxes", [])
        med_text = getattr(self, "popup_c5_medication", None)

        med_vuln_checked = []
        med_prot_checked = []
        for cb in med_cbs:
            if cb.isChecked():
                key = cb.property("c5_key")
                if key in med_vuln_items:
                    med_vuln_checked.append(med_vuln_items[key])
                elif key in med_prot_items:
                    med_prot_checked.append(med_prot_items[key])

        if med_text:
            sentences = []
            if med_vuln_checked:
                sentences.append(f"There are concerns regarding medication adherence, including {join_items(med_vuln_checked)}.")
            if med_prot_checked:
                if sentences:
                    sentences.append(f"However, {subj} demonstrates {join_items(med_prot_checked)}.")
                else:
                    sentences.append(f"{Subj} demonstrates good medication adherence, with {join_items(med_prot_checked)}.")
            if sentences:
                med_text.setPlainText(" ".join(sentences))
            else:
                med_text.clear()

        # ============ 2. ENGAGEMENT WITH SERVICES ============
        eng_vuln_items = {
            "eng_disengaged": "disengagement from services",
            "eng_misses_appts": "missing appointments",
            "eng_poor_attendance": "poor attendance",
            "eng_avoids_reviews": "avoiding reviews",
        }
        eng_prot_items = {
            "eng_actively_engages": "active engagement with services",
        }

        eng_cbs = getattr(self, "popup_c5_engagement_checkboxes", [])
        eng_text = getattr(self, "popup_c5_engagement", None)

        eng_vuln_checked = []
        eng_prot_checked = []
        for cb in eng_cbs:
            if cb.isChecked():
                key = cb.property("c5_key")
                if key in eng_vuln_items:
                    eng_vuln_checked.append(eng_vuln_items[key])
                elif key in eng_prot_items:
                    eng_prot_checked.append(eng_prot_items[key])

        if eng_text:
            sentences = []
            if eng_vuln_checked:
                sentences.append(f"{Subj} has demonstrated poor engagement, with {join_items(eng_vuln_checked)}.")
            if eng_prot_checked:
                if sentences:
                    sentences.append(f"Nevertheless, {subj} now shows {join_items(eng_prot_checked)}.")
                else:
                    sentences.append(f"{Subj} demonstrates {join_items(eng_prot_checked)}.")
            if sentences:
                eng_text.setPlainText(" ".join(sentences))
            else:
                eng_text.clear()

        # ============ 3. COMPLIANCE WITH CONDITIONS ============
        cmp_vuln_items = {
            "cmp_breaches": "breaching conditions",
            "cmp_absconded": "absconding",
            "cmp_recalled": "requiring recall",
            "cmp_only_coerced": "only complying under coercion",
            "cmp_resists_monitoring": "resisting monitoring",
        }
        cmp_prot_items = {
            "cmp_accepts_conditions": "accepting conditions",
        }

        cmp_cbs = getattr(self, "popup_c5_compliance_checkboxes", [])
        cmp_text = getattr(self, "popup_c5_compliance", None)

        cmp_vuln_checked = []
        cmp_prot_checked = []
        for cb in cmp_cbs:
            if cb.isChecked():
                key = cb.property("c5_key")
                if key in cmp_vuln_items:
                    cmp_vuln_checked.append(cmp_vuln_items[key])
                elif key in cmp_prot_items:
                    cmp_prot_checked.append(cmp_prot_items[key])

        if cmp_text:
            sentences = []
            if cmp_vuln_checked:
                sentences.append(f"There is a history of problems with compliance, including {join_items(cmp_vuln_checked)}.")
            if cmp_prot_checked:
                if sentences:
                    sentences.append(f"However, {subj} now demonstrates {join_items(cmp_prot_checked)}.")
                else:
                    sentences.append(f"{Subj} demonstrates {join_items(cmp_prot_checked)}.")
            if sentences:
                cmp_text.setPlainText(" ".join(sentences))
            else:
                cmp_text.clear()

        # ============ 4. PATTERN OVER TIME ============
        pat_vuln_items = {
            "pat_repeated_disengage": "repeated disengagement",
            "pat_history_non_compliance": "a history of non-compliance",
            "pat_cycle_relapse": "a cycle of engagement, discharge, disengagement, and relapse",
        }
        pat_prot_items = {
            "pat_sustained_adherence": "sustained adherence over time",
        }

        pat_cbs = getattr(self, "popup_c5_pattern_checkboxes", [])
        pat_text = getattr(self, "popup_c5_pattern", None)

        pat_vuln_checked = []
        pat_prot_checked = []
        for cb in pat_cbs:
            if cb.isChecked():
                key = cb.property("c5_key")
                if key in pat_vuln_items:
                    pat_vuln_checked.append(pat_vuln_items[key])
                elif key in pat_prot_items:
                    pat_prot_checked.append(pat_prot_items[key])

        if pat_text:
            sentences = []
            if pat_vuln_checked:
                sentences.append(f"Over time, there is a pattern of {join_items(pat_vuln_checked)}.")
            if pat_prot_checked:
                if sentences:
                    sentences.append(f"However, more recently {subj} has shown {join_items(pat_prot_checked)}.")
                else:
                    sentences.append(f"{Subj} has demonstrated {join_items(pat_prot_checked)}.")
            if sentences:
                pat_text.setPlainText(" ".join(sentences))
            else:
                pat_text.clear()

        # ============ 5. TREATMENT RESPONSIVENESS ============
        rsp_vuln_items = {
            "rsp_treatment_resistant": "treatment resistance",
            "rsp_no_improvement": "no improvement with treatment",
        }
        rsp_prot_items = {
            "rsp_responds_well": "responding well to treatment",
            "rsp_benefits_therapy": "benefiting from psychological therapy",
        }

        rsp_cbs = getattr(self, "popup_c5_response_checkboxes", [])
        rsp_text = getattr(self, "popup_c5_response", None)

        rsp_vuln_checked = []
        rsp_prot_checked = []
        for cb in rsp_cbs:
            if cb.isChecked():
                key = cb.property("c5_key")
                if key in rsp_vuln_items:
                    rsp_vuln_checked.append(rsp_vuln_items[key])
                elif key in rsp_prot_items:
                    rsp_prot_checked.append(rsp_prot_items[key])

        if rsp_text:
            sentences = []
            if rsp_vuln_checked:
                sentences.append(f"There are concerns about treatment responsiveness, including {join_items(rsp_vuln_checked)}.")
            if rsp_prot_checked:
                if sentences:
                    sentences.append(f"Nevertheless, {subj} shows positive signs such as {join_items(rsp_prot_checked)}.")
                else:
                    sentences.append(f"{Subj} shows good treatment responsiveness, {join_items(rsp_prot_checked)}.")
            if sentences:
                rsp_text.setPlainText(" ".join(sentences))
            else:
                rsp_text.clear()

    def _build_hcr_r1_popup(self, key: str, code: str, description: str, subsections: list):
        """Build R1 (Professional Services and Plans) popup with tick boxes."""
        container, layout = self._create_popup_container(key)

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Presence/Relevance scoring (standard pattern)
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)

        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            presence_row.addWidget(rb)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)

        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        for rb in [relevance_low, relevance_mod, relevance_high]:
            relevance_row.addWidget(rb)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # === 1. PLAN PRESENCE/CLARITY ===
        plan_container = QFrame()
        plan_container.setStyleSheet("QFrame { background: rgba(219, 234, 254, 0.5); border: 2px solid rgba(59, 130, 246, 0.4); border-radius: 10px; }")
        plan_layout = QVBoxLayout(plan_container)
        plan_layout.setContentsMargins(12, 10, 12, 10)
        plan_layout.setSpacing(8)

        plan_lbl = QLabel("Plan Presence & Clarity")
        plan_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        plan_layout.addWidget(plan_lbl)

        plan_text = QTextEdit()
        plan_text.setPlaceholderText("Describe presence and clarity of care/risk management plan...")
        plan_text.setStyleSheet(input_style + " background: white;")
        plan_text.setMinimumHeight(58)
        plan_layout.addWidget(plan_text)
        plan_layout.addWidget(DragResizeBar(plan_text))
        text_widgets.append(plan_text)
        subsection_data.append(("Plan Presence & Clarity", plan_text))
        setattr(self, "popup_r1_plan", plan_text)

        if CollapsibleSection:
            plan_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            plan_section.set_content_height(140)
            plan_section._min_height = 80
            plan_section._max_height = 220

            plan_content = QWidget()
            plan_content.setStyleSheet("QWidget { background: rgba(239, 246, 255, 0.95); border: 1px solid rgba(59, 130, 246, 0.2); border-top: none; border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; color: #374151; padding: 2px; }")
            plan_cb_layout = QVBoxLayout(plan_content)
            plan_cb_layout.setContentsMargins(10, 6, 10, 6)
            plan_cb_layout.setSpacing(3)

            R1_PLAN_ITEMS = [
                ("pln_clear_plan", "Clear care plan in place"),
                ("pln_risk_plan", "Risk management plan documented"),
                ("pln_no_plan", "No clear aftercare plan"),
                ("pln_incomplete", "Discharge planning incomplete"),
                ("pln_generic", "Generic/vague plan only"),
            ]

            r1_plan_checkboxes = []
            for item_key, item_label in R1_PLAN_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("r1_key", item_key)
                cb.setProperty("r1_category", "plan")
                cb.stateChanged.connect(self._update_r1_services_narrative)
                plan_cb_layout.addWidget(cb)
                r1_plan_checkboxes.append(cb)

            plan_section.set_content(plan_content)
            plan_layout.addWidget(plan_section)
            setattr(self, "popup_r1_plan_checkboxes", r1_plan_checkboxes)

        subsections_target_layout.addWidget(plan_container)

        # === 2. SERVICE ADEQUACY ===
        service_container = QFrame()
        service_container.setStyleSheet("QFrame { background: rgba(254, 243, 199, 0.5); border: 2px solid rgba(217, 119, 6, 0.4); border-radius: 10px; }")
        service_layout = QVBoxLayout(service_container)
        service_layout.setContentsMargins(12, 10, 12, 10)
        service_layout.setSpacing(8)

        service_lbl = QLabel("Service Intensity & Adequacy")
        service_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #b45309; background: transparent; border: none;")
        service_layout.addWidget(service_lbl)

        service_text = QTextEdit()
        service_text.setPlaceholderText("Describe adequacy and intensity of services relative to risk level...")
        service_text.setStyleSheet(input_style + " background: white;")
        service_text.setMinimumHeight(58)
        service_layout.addWidget(service_text)
        service_layout.addWidget(DragResizeBar(service_text))
        text_widgets.append(service_text)
        subsection_data.append(("Service Intensity & Adequacy", service_text))
        setattr(self, "popup_r1_service", service_text)

        if CollapsibleSection:
            service_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            service_section.set_content_height(120)

            service_content = QWidget()
            service_content.setStyleSheet("QWidget { background: rgba(254, 252, 232, 0.95); border: 1px solid rgba(217, 119, 6, 0.2); border-top: none; border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; color: #374151; padding: 2px; }")
            service_cb_layout = QVBoxLayout(service_content)
            service_cb_layout.setContentsMargins(10, 6, 10, 6)
            service_cb_layout.setSpacing(3)

            R1_SERVICE_ITEMS = [
                ("svc_appropriate", "Services appropriate to risk level"),
                ("svc_insufficient", "Insufficient support for risk level"),
                ("svc_limited", "Limited community input planned"),
                ("svc_mismatch", "Risk-service mismatch"),
            ]

            r1_service_checkboxes = []
            for item_key, item_label in R1_SERVICE_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("r1_key", item_key)
                cb.setProperty("r1_category", "service")
                cb.stateChanged.connect(self._update_r1_services_narrative)
                service_cb_layout.addWidget(cb)
                r1_service_checkboxes.append(cb)

            service_section.set_content(service_content)
            service_layout.addWidget(service_section)
            setattr(self, "popup_r1_service_checkboxes", r1_service_checkboxes)

        subsections_target_layout.addWidget(service_container)

        # === 3. TRANSITIONS/CONTINUITY ===
        transitions_container = QFrame()
        transitions_container.setStyleSheet("QFrame { background: rgba(254, 226, 226, 0.5); border: 2px solid rgba(220, 38, 38, 0.4); border-radius: 10px; }")
        transitions_layout = QVBoxLayout(transitions_container)
        transitions_layout.setContentsMargins(12, 10, 12, 10)
        transitions_layout.setSpacing(8)

        transitions_lbl = QLabel("Transitions & Continuity")
        transitions_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        transitions_layout.addWidget(transitions_lbl)

        transitions_text = QTextEdit()
        transitions_text.setPlaceholderText("Describe care transitions, gaps, handover arrangements...")
        transitions_text.setStyleSheet(input_style + " background: white;")
        transitions_text.setMinimumHeight(58)
        transitions_layout.addWidget(transitions_text)
        transitions_layout.addWidget(DragResizeBar(transitions_text))
        text_widgets.append(transitions_text)
        subsection_data.append(("Transitions & Continuity", transitions_text))
        setattr(self, "popup_r1_transitions", transitions_text)

        if CollapsibleSection:
            transitions_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            transitions_section.set_content_height(120)

            transitions_content = QWidget()
            transitions_content.setStyleSheet("QWidget { background: rgba(254, 242, 242, 0.95); border: 1px solid rgba(220, 38, 38, 0.2); border-top: none; border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; color: #374151; padding: 2px; }")
            transitions_cb_layout = QVBoxLayout(transitions_content)
            transitions_cb_layout.setContentsMargins(10, 6, 10, 6)
            transitions_cb_layout.setSpacing(3)

            R1_TRANSITIONS_ITEMS = [
                ("trn_awaiting", "Awaiting allocation"),
                ("trn_waiting_list", "On waiting list"),
                ("trn_no_followup", "No confirmed follow-up"),
                ("trn_gap", "Gap in care likely"),
                ("trn_timely", "Timely follow-up arranged (protective)"),
            ]

            r1_transitions_checkboxes = []
            for item_key, item_label in R1_TRANSITIONS_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("r1_key", item_key)
                cb.setProperty("r1_category", "transitions")
                cb.stateChanged.connect(self._update_r1_services_narrative)
                transitions_cb_layout.addWidget(cb)
                r1_transitions_checkboxes.append(cb)

            transitions_section.set_content(transitions_content)
            transitions_layout.addWidget(transitions_section)
            setattr(self, "popup_r1_transitions_checkboxes", r1_transitions_checkboxes)

        subsections_target_layout.addWidget(transitions_container)

        # === 4. CONTINGENCY PLANNING ===
        contingency_container = QFrame()
        contingency_container.setStyleSheet("QFrame { background: rgba(220, 252, 231, 0.5); border: 2px solid rgba(22, 163, 74, 0.4); border-radius: 10px; }")
        contingency_layout = QVBoxLayout(contingency_container)
        contingency_layout.setContentsMargins(12, 10, 12, 10)
        contingency_layout.setSpacing(8)

        contingency_lbl = QLabel("Contingency & Escalation Planning")
        contingency_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #15803d; background: transparent; border: none;")
        contingency_layout.addWidget(contingency_lbl)

        contingency_text = QTextEdit()
        contingency_text.setPlaceholderText("Describe crisis planning, early warning signs, escalation pathways...")
        contingency_text.setStyleSheet(input_style + " background: white;")
        contingency_text.setMinimumHeight(58)
        contingency_layout.addWidget(contingency_text)
        contingency_layout.addWidget(DragResizeBar(contingency_text))
        text_widgets.append(contingency_text)
        subsection_data.append(("Contingency & Escalation Planning", contingency_text))
        setattr(self, "popup_r1_contingency", contingency_text)

        if CollapsibleSection:
            contingency_section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
            contingency_section.set_content_height(140)

            contingency_content = QWidget()
            contingency_content.setStyleSheet("QWidget { background: rgba(240, 253, 244, 0.95); border: 1px solid rgba(22, 163, 74, 0.2); border-top: none; border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; color: #374151; padding: 2px; }")
            contingency_cb_layout = QVBoxLayout(contingency_content)
            contingency_cb_layout.setContentsMargins(10, 6, 10, 6)
            contingency_cb_layout.setSpacing(3)

            R1_CONTINGENCY_ITEMS = [
                ("cnt_crisis_plan", "Crisis plan in place"),
                ("cnt_warning_signs", "Early warning signs documented"),
                ("cnt_escalation", "Clear escalation pathway"),
                ("cnt_no_crisis", "No crisis plan"),
                ("cnt_no_escalation", "Unclear escalation pathway"),
            ]

            r1_contingency_checkboxes = []
            for item_key, item_label in R1_CONTINGENCY_ITEMS:
                cb = QCheckBox(item_label)
                cb.setProperty("r1_key", item_key)
                cb.setProperty("r1_category", "contingency")
                cb.stateChanged.connect(self._update_r1_services_narrative)
                contingency_cb_layout.addWidget(cb)
                r1_contingency_checkboxes.append(cb)

            contingency_section.set_content(contingency_content)
            contingency_layout.addWidget(contingency_section)
            setattr(self, "popup_r1_contingency_checkboxes", r1_contingency_checkboxes)

        subsections_target_layout.addWidget(contingency_container)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        # === IMPORTED DATA ===
        if CollapsibleSection:
            imported_section = CollapsibleSection("Imported Data", start_collapsed=True)
            imported_section.set_content_height(200)
            imported_section._min_height = 80
            imported_section._max_height = 400

            imported_content = QWidget()
            imported_content.setStyleSheet("QWidget { background: rgba(255, 248, 220, 0.95); border: 1px solid rgba(180, 150, 50, 0.4); border-top: none; border-radius: 0 0 10px 10px; }")
            imported_content_layout = QVBoxLayout(imported_content)
            imported_content_layout.setContentsMargins(6, 6, 6, 6)

            imported_scroll = QScrollArea()
            imported_scroll.setWidgetResizable(True)
            imported_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
            imported_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

            imported_entries_container = QWidget()
            imported_entries_container.setStyleSheet("background: transparent;")
            imported_entries_layout = QVBoxLayout(imported_entries_container)
            imported_entries_layout.setContentsMargins(2, 2, 2, 2)
            imported_entries_layout.setSpacing(6)
            imported_entries_layout.setAlignment(Qt.AlignmentFlag.AlignTop)

            imported_scroll.setWidget(imported_entries_container)
            imported_content_layout.addWidget(imported_scroll)

            imported_section.set_content(imported_content)
            layout.addWidget(imported_section)

            setattr(self, f"popup_{key}_imported_section", imported_section)
            setattr(self, f"popup_{key}_imported_entries_layout", imported_entries_layout)
            setattr(self, f"popup_{key}_imported_checkboxes", [])

        layout.addStretch()

        # Store references
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_presence_no", presence_no)
        setattr(self, f"popup_{key}_presence_partial", presence_partial)
        setattr(self, f"popup_{key}_presence_yes", presence_yes)
        setattr(self, f"popup_{key}_presence_omit", presence_omit)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_relevance_low", relevance_low)
        setattr(self, f"popup_{key}_relevance_mod", relevance_mod)
        setattr(self, f"popup_{key}_relevance_high", relevance_high)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]
            presence_val = ""
            if presence_no.isChecked():
                presence_val = "No"
            elif presence_partial.isChecked():
                presence_val = "Partial/Possible"
            elif presence_yes.isChecked():
                presence_val = "Present"
            elif presence_omit.isChecked():
                presence_val = "Omitted"
            if presence_val:
                parts.append(f"Presence: {presence_val}")

            relevance_val = ""
            if relevance_low.isChecked():
                relevance_val = "Low relevance"
            elif relevance_mod.isChecked():
                relevance_val = "Moderate relevance"
            elif relevance_high.isChecked():
                relevance_val = "High relevance"
            if relevance_val:
                parts.append(f"Relevance: {relevance_val}")
            parts.append("")

            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}:")
                    parts.append(text)
                    parts.append("")
            return "\n".join(parts)

        setattr(self, f"popup_{key}_generate", generate)

    def _update_r1_services_narrative(self):
        """Update R1 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss = 'he', 'him', 'his'
        else:
            subj, obj, poss = 'she', 'her', 'her'
        Subj = subj.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. CARE PLANNING ============
        pln_prot_items = {"pln_clear_plan": "a clear care plan", "pln_risk_plan": "a documented risk management plan"}
        pln_vuln_items = {"pln_no_plan": "no clear aftercare plan", "pln_incomplete": "incomplete discharge planning", "pln_generic": "only a generic/vague plan"}

        pln_cbs = getattr(self, "popup_r1_plan_checkboxes", [])
        pln_text = getattr(self, "popup_r1_plan", None)

        pln_prot, pln_vuln = [], []
        for cb in pln_cbs:
            if cb.isChecked():
                key = cb.property("r1_key")
                if key in pln_prot_items:
                    pln_prot.append(pln_prot_items[key])
                elif key in pln_vuln_items:
                    pln_vuln.append(pln_vuln_items[key])

        if pln_text:
            sentences = []
            if pln_prot:
                sentences.append(f"There is {join_items(pln_prot)} in place.")
            if pln_vuln:
                if sentences:
                    sentences.append(f"However, concerns include {join_items(pln_vuln)}.")
                else:
                    sentences.append(f"There are concerns regarding care planning, including {join_items(pln_vuln)}.")
            pln_text.setPlainText(" ".join(sentences)) if sentences else pln_text.clear()

        # ============ 2. SERVICE ADEQUACY ============
        svc_prot_items = {"svc_appropriate": "services appropriate to risk level"}
        svc_vuln_items = {"svc_insufficient": "insufficient support", "svc_limited": "limited community input", "svc_mismatch": "a risk-service mismatch"}

        svc_cbs = getattr(self, "popup_r1_service_checkboxes", [])
        svc_text = getattr(self, "popup_r1_service", None)

        svc_prot, svc_vuln = [], []
        for cb in svc_cbs:
            if cb.isChecked():
                key = cb.property("r1_key")
                if key in svc_prot_items:
                    svc_prot.append(svc_prot_items[key])
                elif key in svc_vuln_items:
                    svc_vuln.append(svc_vuln_items[key])

        if svc_text:
            sentences = []
            if svc_vuln:
                sentences.append(f"Service adequacy concerns include {join_items(svc_vuln)}.")
            if svc_prot:
                if sentences:
                    sentences.append(f"However, there are {join_items(svc_prot)}.")
                else:
                    sentences.append(f"There are {join_items(svc_prot)}.")
            svc_text.setPlainText(" ".join(sentences)) if sentences else svc_text.clear()

        # ============ 3. TRANSITIONS/CONTINUITY ============
        trn_prot_items = {"trn_timely": "timely follow-up arranged"}
        trn_vuln_items = {"trn_awaiting": "awaiting allocation", "trn_waiting_list": "on a waiting list", "trn_no_followup": "no confirmed follow-up", "trn_gap": "a likely gap in care"}

        trn_cbs = getattr(self, "popup_r1_transitions_checkboxes", [])
        trn_text = getattr(self, "popup_r1_transitions", None)

        trn_prot, trn_vuln = [], []
        for cb in trn_cbs:
            if cb.isChecked():
                key = cb.property("r1_key")
                if key in trn_prot_items:
                    trn_prot.append(trn_prot_items[key])
                elif key in trn_vuln_items:
                    trn_vuln.append(trn_vuln_items[key])

        if trn_text:
            sentences = []
            if trn_vuln:
                sentences.append(f"Transition concerns include {join_items(trn_vuln)}.")
            if trn_prot:
                if sentences:
                    sentences.append(f"Nevertheless, there is {join_items(trn_prot)}.")
                else:
                    sentences.append(f"There is {join_items(trn_prot)}.")
            trn_text.setPlainText(" ".join(sentences)) if sentences else trn_text.clear()

        # ============ 4. CONTINGENCY PLANNING ============
        cnt_prot_items = {"cnt_crisis_plan": "a crisis plan", "cnt_warning_signs": "documented early warning signs", "cnt_escalation": "a clear escalation pathway"}
        cnt_vuln_items = {"cnt_no_crisis": "no crisis plan", "cnt_no_escalation": "an unclear escalation pathway"}

        cnt_cbs = getattr(self, "popup_r1_contingency_checkboxes", [])
        cnt_text = getattr(self, "popup_r1_contingency", None)

        cnt_prot, cnt_vuln = [], []
        for cb in cnt_cbs:
            if cb.isChecked():
                key = cb.property("r1_key")
                if key in cnt_prot_items:
                    cnt_prot.append(cnt_prot_items[key])
                elif key in cnt_vuln_items:
                    cnt_vuln.append(cnt_vuln_items[key])

        if cnt_text:
            sentences = []
            if cnt_prot:
                sentences.append(f"Contingency planning includes {join_items(cnt_prot)}.")
            if cnt_vuln:
                if sentences:
                    sentences.append(f"However, there is {join_items(cnt_vuln)}.")
                else:
                    sentences.append(f"There are concerns including {join_items(cnt_vuln)}.")
            cnt_text.setPlainText(" ".join(sentences)) if sentences else cnt_text.clear()

    def _build_hcr_r2_popup(self, key: str, code: str, description: str, subsections: list):
        """Build R2 (Living Situation) popup with tick boxes."""
        container, layout = self._create_popup_container(key)
        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        # Standard presence/relevance scoring
        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)
        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
            presence_row.addWidget(rb)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)
        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
            relevance_row.addWidget(rb)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        # R2 Categories with checkboxes
        categories = [
            ("Accommodation Stability", "rgba(254, 226, 226, 0.5)", "rgba(220, 38, 38, 0.4)", "#991b1b", "r2_accom", [
                ("accom_unstable", "Unstable housing / NFA"),
                ("accom_temporary", "Temporary accommodation"),
                ("accom_eviction_risk", "At risk of eviction"),
                ("accom_stable", "Stable accommodation (protective)"),
            ]),
            ("Who They Live With", "rgba(254, 243, 199, 0.5)", "rgba(217, 119, 6, 0.4)", "#b45309", "r2_cohab", [
                ("cohab_victim", "Living with/near victim"),
                ("cohab_conflict", "Conflictual family environment"),
                ("cohab_substance_peers", "Living with substance-using peers"),
                ("cohab_supportive", "Supportive household (protective)"),
            ]),
            ("Supervision Level", "rgba(219, 234, 254, 0.5)", "rgba(59, 130, 246, 0.4)", "#1d4ed8", "r2_super", [
                ("super_supported", "Supported/staffed setting"),
                ("super_unsupervised", "Completely unsupervised"),
                ("super_step_down", "Step-down without preparation"),
                ("super_deteriorates", "Deteriorates without support"),
            ]),
            ("Substance Access", "rgba(220, 252, 231, 0.5)", "rgba(22, 163, 74, 0.4)", "#15803d", "r2_subst", [
                ("subst_access", "Easy access to substances"),
                ("subst_peers", "Substance-using peers nearby"),
                ("subst_free", "Substance-free environment (protective)"),
            ]),
        ]

        for cat_name, bg_color, border_color, text_color, cat_key, items in categories:
            cont = QFrame()
            cont.setStyleSheet(f"QFrame {{ background: {bg_color}; border: 2px solid {border_color}; border-radius: 10px; }}")
            cont_layout = QVBoxLayout(cont)
            cont_layout.setContentsMargins(12, 10, 12, 10)
            cont_layout.setSpacing(8)

            lbl = QLabel(cat_name)
            lbl.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {text_color}; background: transparent; border: none;")
            cont_layout.addWidget(lbl)

            txt = QTextEdit()
            txt.setPlaceholderText(f"Describe {cat_name.lower()}...")
            txt.setStyleSheet(input_style + " background: white;")
            txt.setMinimumHeight(58)
            cont_layout.addWidget(txt)
            cont_layout.addWidget(DragResizeBar(txt))
            text_widgets.append(txt)
            subsection_data.append((cat_name, txt))
            setattr(self, f"popup_{cat_key}", txt)

            if CollapsibleSection:
                section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
                section.set_content_height(100)
                content = QWidget()
                content.setStyleSheet("QWidget { background: rgba(255,255,255,0.9); border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; padding: 2px; }")
                cb_layout = QVBoxLayout(content)
                cb_layout.setContentsMargins(10, 6, 10, 6)
                cb_layout.setSpacing(3)

                checkboxes = []
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("r2_key", item_key)
                    cb.setProperty("r2_category", cat_key)
                    cb.stateChanged.connect(self._update_r2_living_narrative)
                    cb_layout.addWidget(cb)
                    checkboxes.append(cb)

                section.set_content(content)
                cont_layout.addWidget(section)
                setattr(self, f"popup_{cat_key}_checkboxes", checkboxes)

            subsections_target_layout.addWidget(cont)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        layout.addStretch()
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]
            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}: {text}")
            return "\n".join(parts)
        setattr(self, f"popup_{key}_generate", generate)

    def _update_r2_living_narrative(self):
        """Update R2 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss = 'he', 'him', 'his'
        else:
            subj, obj, poss = 'she', 'her', 'her'
        Subj = subj.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. ACCOMMODATION STABILITY ============
        accom_vuln = {"accom_unstable": "unstable housing", "accom_temporary": "temporary accommodation", "accom_eviction_risk": "risk of eviction"}
        accom_prot = {"accom_stable": "stable accommodation"}

        accom_cbs = getattr(self, "popup_r2_accom_checkboxes", [])
        accom_text = getattr(self, "popup_r2_accom", None)

        accom_v, accom_p = [], []
        for cb in accom_cbs:
            if cb.isChecked():
                key = cb.property("r2_key")
                if key in accom_vuln:
                    accom_v.append(accom_vuln[key])
                elif key in accom_prot:
                    accom_p.append(accom_prot[key])

        if accom_text:
            sentences = []
            if accom_v:
                sentences.append(f"Accommodation concerns include {join_items(accom_v)}.")
            if accom_p:
                sentences.append(f"{Subj} has {join_items(accom_p)}.")
            accom_text.setPlainText(" ".join(sentences)) if sentences else accom_text.clear()

        # ============ 2. WHO THEY LIVE WITH ============
        cohab_vuln = {"cohab_victim": "living with/near a victim", "cohab_conflict": "a conflictual family environment", "cohab_substance_peers": "substance-using peers"}
        cohab_prot = {"cohab_supportive": "a supportive household"}

        cohab_cbs = getattr(self, "popup_r2_cohab_checkboxes", [])
        cohab_text = getattr(self, "popup_r2_cohab", None)

        cohab_v, cohab_p = [], []
        for cb in cohab_cbs:
            if cb.isChecked():
                key = cb.property("r2_key")
                if key in cohab_vuln:
                    cohab_v.append(cohab_vuln[key])
                elif key in cohab_prot:
                    cohab_p.append(cohab_prot[key])

        if cohab_text:
            sentences = []
            if cohab_v:
                sentences.append(f"Concerning cohabitation factors include {join_items(cohab_v)}.")
            if cohab_p:
                if sentences:
                    sentences.append(f"However, {subj} has {join_items(cohab_p)}.")
                else:
                    sentences.append(f"{Subj} has {join_items(cohab_p)}.")
            cohab_text.setPlainText(" ".join(sentences)) if sentences else cohab_text.clear()

        # ============ 3. SUPERVISION LEVEL ============
        super_items = {
            "super_supported": ("prot", "a supported/staffed setting"),
            "super_unsupervised": ("vuln", "completely unsupervised accommodation"),
            "super_step_down": ("vuln", "a step-down without adequate preparation"),
            "super_deteriorates": ("vuln", "deterioration without support"),
        }

        super_cbs = getattr(self, "popup_r2_super_checkboxes", [])
        super_text = getattr(self, "popup_r2_super", None)

        super_v, super_p = [], []
        for cb in super_cbs:
            if cb.isChecked():
                key = cb.property("r2_key")
                if key in super_items:
                    typ, phrase = super_items[key]
                    if typ == "vuln":
                        super_v.append(phrase)
                    else:
                        super_p.append(phrase)

        if super_text:
            sentences = []
            if super_v:
                sentences.append(f"Supervision concerns include {join_items(super_v)}.")
            if super_p:
                sentences.append(f"{Subj} is in {join_items(super_p)}.")
            super_text.setPlainText(" ".join(sentences)) if sentences else super_text.clear()

        # ============ 4. SUBSTANCE ACCESS ============
        subst_vuln = {"subst_access": "easy access to substances", "subst_peers": "substance-using peers nearby"}
        subst_prot = {"subst_free": "a substance-free environment"}

        subst_cbs = getattr(self, "popup_r2_subst_checkboxes", [])
        subst_text = getattr(self, "popup_r2_subst", None)

        subst_v, subst_p = [], []
        for cb in subst_cbs:
            if cb.isChecked():
                key = cb.property("r2_key")
                if key in subst_vuln:
                    subst_v.append(subst_vuln[key])
                elif key in subst_prot:
                    subst_p.append(subst_prot[key])

        if subst_text:
            sentences = []
            if subst_v:
                sentences.append(f"Substance-related concerns include {join_items(subst_v)}.")
            if subst_p:
                if sentences:
                    sentences.append(f"However, {subj} is in {join_items(subst_p)}.")
                else:
                    sentences.append(f"{Subj} is in {join_items(subst_p)}.")
            subst_text.setPlainText(" ".join(sentences)) if sentences else subst_text.clear()

    def _build_hcr_r3_popup(self, key: str, code: str, description: str, subsections: list):
        """Build R3 (Personal Support) popup with tick boxes."""
        container, layout = self._create_popup_container(key)
        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)
        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
            presence_row.addWidget(rb)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)
        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
            relevance_row.addWidget(rb)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        categories = [
            ("Supportive Relationships", "rgba(220, 252, 231, 0.5)", "#15803d", "r3_support", [
                ("sup_family", "Supportive family"),
                ("sup_partner", "Supportive partner"),
                ("sup_regular_contact", "Regular contact with supports"),
                ("sup_crisis_help", "Support available in crisis"),
            ]),
            ("Isolation/Weak Support", "rgba(254, 226, 226, 0.5)", "#991b1b", "r3_isolation", [
                ("iso_limited", "Limited social support"),
                ("iso_estranged", "Estranged from family"),
                ("iso_lives_alone", "Lives alone with limited contact"),
                ("iso_superficial", "Superficial/unreliable contacts only"),
            ]),
            ("Conflict Within Network", "rgba(254, 243, 199, 0.5)", "#b45309", "r3_conflict", [
                ("con_interpersonal", "High interpersonal conflict"),
                ("con_volatile", "Volatile relationships"),
                ("con_domestic", "Domestic conflict"),
            ]),
            ("Antisocial Peers", "rgba(243, 232, 255, 0.5)", "#7e22ce", "r3_peers", [
                ("peer_antisocial", "Mixes with antisocial peers"),
                ("peer_substance", "Substance-using peers"),
                ("peer_negative", "Negative peer influence"),
            ]),
        ]

        for cat_name, bg_color, text_color, cat_key, items in categories:
            cont = QFrame()
            cont.setStyleSheet(f"QFrame {{ background: {bg_color}; border: 2px solid {text_color}40; border-radius: 10px; }}")
            cont_layout = QVBoxLayout(cont)
            cont_layout.setContentsMargins(12, 10, 12, 10)
            cont_layout.setSpacing(8)

            lbl = QLabel(cat_name)
            lbl.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {text_color}; background: transparent; border: none;")
            cont_layout.addWidget(lbl)

            txt = QTextEdit()
            txt.setPlaceholderText(f"Describe {cat_name.lower()}...")
            txt.setStyleSheet(input_style + " background: white;")
            txt.setMinimumHeight(58)
            cont_layout.addWidget(txt)
            cont_layout.addWidget(DragResizeBar(txt))
            text_widgets.append(txt)
            subsection_data.append((cat_name, txt))
            setattr(self, f"popup_{cat_key}", txt)

            if CollapsibleSection:
                section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
                section.set_content_height(100)
                content = QWidget()
                content.setStyleSheet("QWidget { background: rgba(255,255,255,0.9); border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; padding: 2px; }")
                cb_layout = QVBoxLayout(content)
                cb_layout.setContentsMargins(10, 6, 10, 6)
                cb_layout.setSpacing(3)

                checkboxes = []
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("r3_key", item_key)
                    cb.setProperty("r3_category", cat_key)
                    cb.stateChanged.connect(self._update_r3_support_narrative)
                    cb_layout.addWidget(cb)
                    checkboxes.append(cb)

                section.set_content(content)
                cont_layout.addWidget(section)
                setattr(self, f"popup_{cat_key}_checkboxes", checkboxes)

            subsections_target_layout.addWidget(cont)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        layout.addStretch()
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]
            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}: {text}")
            return "\n".join(parts)
        setattr(self, f"popup_{key}_generate", generate)

    def _update_r3_support_narrative(self):
        """Update R3 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss = 'he', 'him', 'his'
        else:
            subj, obj, poss = 'she', 'her', 'her'
        Subj = subj.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. SUPPORTIVE RELATIONSHIPS (all protective) ============
        sup_items = {"sup_family": "supportive family", "sup_partner": "a supportive partner", "sup_regular_contact": "regular contact with supports", "sup_crisis_help": "support available in crisis"}

        sup_cbs = getattr(self, "popup_r3_support_checkboxes", [])
        sup_text = getattr(self, "popup_r3_support", None)

        sup_checked = []
        for cb in sup_cbs:
            if cb.isChecked():
                key = cb.property("r3_key")
                if key in sup_items:
                    sup_checked.append(sup_items[key])

        if sup_text:
            if sup_checked:
                sup_text.setPlainText(f"Protective factors include {join_items(sup_checked)}.")
            else:
                sup_text.clear()

        # ============ 2. ISOLATION/WEAK SUPPORT (all vulnerability) ============
        iso_items = {"iso_limited": "limited social support", "iso_estranged": "being estranged from family", "iso_lives_alone": "living alone with limited contact", "iso_superficial": "superficial or unreliable contacts"}

        iso_cbs = getattr(self, "popup_r3_isolation_checkboxes", [])
        iso_text = getattr(self, "popup_r3_isolation", None)

        iso_checked = []
        for cb in iso_cbs:
            if cb.isChecked():
                key = cb.property("r3_key")
                if key in iso_items:
                    iso_checked.append(iso_items[key])

        if iso_text:
            if iso_checked:
                iso_text.setPlainText(f"{Subj} experiences social isolation, including {join_items(iso_checked)}.")
            else:
                iso_text.clear()

        # ============ 3. CONFLICT WITHIN NETWORK (all vulnerability) ============
        con_items = {"con_interpersonal": "high interpersonal conflict", "con_volatile": "volatile relationships", "con_domestic": "domestic conflict"}

        con_cbs = getattr(self, "popup_r3_conflict_checkboxes", [])
        con_text = getattr(self, "popup_r3_conflict", None)

        con_checked = []
        for cb in con_cbs:
            if cb.isChecked():
                key = cb.property("r3_key")
                if key in con_items:
                    con_checked.append(con_items[key])

        if con_text:
            if con_checked:
                con_text.setPlainText(f"There is conflict within {poss} support network, including {join_items(con_checked)}.")
            else:
                con_text.clear()

        # ============ 4. ANTISOCIAL PEERS (all vulnerability) ============
        peer_items = {"peer_antisocial": "mixing with antisocial peers", "peer_substance": "substance-using peers", "peer_negative": "negative peer influence"}

        peer_cbs = getattr(self, "popup_r3_peers_checkboxes", [])
        peer_text = getattr(self, "popup_r3_peers", None)

        peer_checked = []
        for cb in peer_cbs:
            if cb.isChecked():
                key = cb.property("r3_key")
                if key in peer_items:
                    peer_checked.append(peer_items[key])

        if peer_text:
            if peer_checked:
                peer_text.setPlainText(f"There are concerns about peer associations, including {join_items(peer_checked)}.")
            else:
                peer_text.clear()

    def _build_hcr_r4_popup(self, key: str, code: str, description: str, subsections: list):
        """Build R4 (Treatment/Supervision Compliance) popup with tick boxes."""
        container, layout = self._create_popup_container(key)
        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)
        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
            presence_row.addWidget(rb)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)
        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
            relevance_row.addWidget(rb)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        categories = [
            ("Medication Adherence", "rgba(254, 226, 226, 0.5)", "#991b1b", "r4_med", [
                ("med_likely_stop", "Likely to stop medication"),
                ("med_likely_refuse", "Likely to refuse medication"),
                ("med_history_noncompliance", "History of medication non-compliance"),
                ("med_likely_comply", "Likely to comply with medication (protective)"),
            ]),
            ("Attendance/Engagement", "rgba(254, 243, 199, 0.5)", "#b45309", "r4_attend", [
                ("att_likely_miss", "Likely to miss appointments"),
                ("att_likely_disengage", "Likely to disengage from services"),
                ("att_history_dna", "History of DNAs"),
                ("att_likely_engage", "Likely to engage (protective)"),
            ]),
            ("Supervision Compliance", "rgba(243, 232, 255, 0.5)", "#7e22ce", "r4_super", [
                ("sup_likely_breach", "Likely to breach conditions"),
                ("sup_history_breach", "History of breaching conditions"),
                ("sup_only_coerced", "Only compliant under coercion"),
                ("sup_likely_accept", "Likely to accept supervision (protective)"),
            ]),
            ("Response to Enforcement", "rgba(219, 234, 254, 0.5)", "#1d4ed8", "r4_enforce", [
                ("enf_hostile", "Becomes hostile when supervised"),
                ("enf_resists", "Resists monitoring"),
                ("enf_escalates", "Escalates when challenged"),
                ("enf_accepts", "Accepts enforcement constructively (protective)"),
            ]),
        ]

        for cat_name, bg_color, text_color, cat_key, items in categories:
            cont = QFrame()
            cont.setStyleSheet(f"QFrame {{ background: {bg_color}; border: 2px solid {text_color}40; border-radius: 10px; }}")
            cont_layout = QVBoxLayout(cont)
            cont_layout.setContentsMargins(12, 10, 12, 10)
            cont_layout.setSpacing(8)

            lbl = QLabel(cat_name)
            lbl.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {text_color}; background: transparent; border: none;")
            cont_layout.addWidget(lbl)

            txt = QTextEdit()
            txt.setPlaceholderText(f"Describe likely {cat_name.lower()}...")
            txt.setStyleSheet(input_style + " background: white;")
            txt.setMinimumHeight(58)
            cont_layout.addWidget(txt)
            cont_layout.addWidget(DragResizeBar(txt))
            text_widgets.append(txt)
            subsection_data.append((cat_name, txt))
            setattr(self, f"popup_{cat_key}", txt)

            if CollapsibleSection:
                section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
                section.set_content_height(100)
                content = QWidget()
                content.setStyleSheet("QWidget { background: rgba(255,255,255,0.9); border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; padding: 2px; }")
                cb_layout = QVBoxLayout(content)
                cb_layout.setContentsMargins(10, 6, 10, 6)
                cb_layout.setSpacing(3)

                checkboxes = []
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("r4_key", item_key)
                    cb.setProperty("r4_category", cat_key)
                    cb.stateChanged.connect(self._update_r4_compliance_narrative)
                    cb_layout.addWidget(cb)
                    checkboxes.append(cb)

                section.set_content(content)
                cont_layout.addWidget(section)
                setattr(self, f"popup_{cat_key}_checkboxes", checkboxes)

            subsections_target_layout.addWidget(cont)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        layout.addStretch()
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]
            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}: {text}")
            return "\n".join(parts)
        setattr(self, f"popup_{key}_generate", generate)

    def _update_r4_compliance_narrative(self):
        """Update R4 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss = 'he', 'him', 'his'
        else:
            subj, obj, poss = 'she', 'her', 'her'
        Subj = subj.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. MEDICATION ADHERENCE ============
        med_vuln = {"med_likely_stop": "being likely to stop medication", "med_likely_refuse": "being likely to refuse medication", "med_history_noncompliance": "a history of non-compliance"}
        med_prot = {"med_likely_comply": "likely compliance with medication"}

        med_cbs = getattr(self, "popup_r4_med_checkboxes", [])
        med_text = getattr(self, "popup_r4_med", None)

        med_v, med_p = [], []
        for cb in med_cbs:
            if cb.isChecked():
                key = cb.property("r4_key")
                if key in med_vuln:
                    med_v.append(med_vuln[key])
                elif key in med_prot:
                    med_p.append(med_prot[key])

        if med_text:
            sentences = []
            if med_v:
                sentences.append(f"Anticipated medication adherence concerns include {join_items(med_v)}.")
            if med_p:
                if sentences:
                    sentences.append(f"However, {subj} shows {join_items(med_p)}.")
                else:
                    sentences.append(f"{Subj} shows {join_items(med_p)}.")
            med_text.setPlainText(" ".join(sentences)) if sentences else med_text.clear()

        # ============ 2. ATTENDANCE ============
        att_vuln = {"att_likely_miss": "being likely to miss appointments", "att_likely_disengage": "being likely to disengage", "att_history_dna": "a history of DNAs"}
        att_prot = {"att_likely_engage": "likely engagement with services"}

        att_cbs = getattr(self, "popup_r4_attend_checkboxes", [])
        att_text = getattr(self, "popup_r4_attend", None)

        att_v, att_p = [], []
        for cb in att_cbs:
            if cb.isChecked():
                key = cb.property("r4_key")
                if key in att_vuln:
                    att_v.append(att_vuln[key])
                elif key in att_prot:
                    att_p.append(att_prot[key])

        if att_text:
            sentences = []
            if att_v:
                sentences.append(f"Anticipated attendance concerns include {join_items(att_v)}.")
            if att_p:
                if sentences:
                    sentences.append(f"Nevertheless, {subj} shows {join_items(att_p)}.")
                else:
                    sentences.append(f"{Subj} shows {join_items(att_p)}.")
            att_text.setPlainText(" ".join(sentences)) if sentences else att_text.clear()

        # ============ 3. SUPERVISION ============
        sup_vuln = {"sup_likely_breach": "being likely to breach conditions", "sup_history_breach": "a history of breaches", "sup_only_coerced": "only complying under coercion"}
        sup_prot = {"sup_likely_accept": "likely acceptance of supervision"}

        sup_cbs = getattr(self, "popup_r4_super_checkboxes", [])
        sup_text = getattr(self, "popup_r4_super", None)

        sup_v, sup_p = [], []
        for cb in sup_cbs:
            if cb.isChecked():
                key = cb.property("r4_key")
                if key in sup_vuln:
                    sup_v.append(sup_vuln[key])
                elif key in sup_prot:
                    sup_p.append(sup_prot[key])

        if sup_text:
            sentences = []
            if sup_v:
                sentences.append(f"Anticipated supervision concerns include {join_items(sup_v)}.")
            if sup_p:
                if sentences:
                    sentences.append(f"However, {subj} shows {join_items(sup_p)}.")
                else:
                    sentences.append(f"{Subj} shows {join_items(sup_p)}.")
            sup_text.setPlainText(" ".join(sentences)) if sentences else sup_text.clear()

        # ============ 4. RESPONSE TO ENFORCEMENT ============
        enf_vuln = {"enf_hostile": "becoming hostile when supervised", "enf_resists": "resisting monitoring", "enf_escalates": "escalating when challenged"}
        enf_prot = {"enf_accepts": "accepting enforcement constructively"}

        enf_cbs = getattr(self, "popup_r4_enforce_checkboxes", [])
        enf_text = getattr(self, "popup_r4_enforce", None)

        enf_v, enf_p = [], []
        for cb in enf_cbs:
            if cb.isChecked():
                key = cb.property("r4_key")
                if key in enf_vuln:
                    enf_v.append(enf_vuln[key])
                elif key in enf_prot:
                    enf_p.append(enf_prot[key])

        if enf_text:
            sentences = []
            if enf_v:
                sentences.append(f"Anticipated response to enforcement concerns include {join_items(enf_v)}.")
            if enf_p:
                if sentences:
                    sentences.append(f"Nevertheless, {subj} is capable of {join_items(enf_p)}.")
                else:
                    sentences.append(f"{Subj} is capable of {join_items(enf_p)}.")
            enf_text.setPlainText(" ".join(sentences)) if sentences else enf_text.clear()

    def _build_hcr_r5_popup(self, key: str, code: str, description: str, subsections: list):
        """Build R5 (Stress or Coping) popup with tick boxes."""
        container, layout = self._create_popup_container(key)
        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 16px; }"

        desc_lbl = QLabel(f"{code}: {description}")
        desc_lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        desc_lbl.setWordWrap(True)
        layout.addWidget(desc_lbl)

        presence_lbl = QLabel("Presence:")
        presence_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 12px;")
        layout.addWidget(presence_lbl)
        presence_row = QHBoxLayout()
        presence_group = QButtonGroup(self)
        presence_no = QRadioButton("No (N)")
        presence_partial = QRadioButton("Partial (P)")
        presence_yes = QRadioButton("Yes (Y)")
        presence_omit = QRadioButton("Omit")
        for rb in [presence_no, presence_partial, presence_yes, presence_omit]:
            rb.setStyleSheet(radio_style)
            presence_row.addWidget(rb)
        presence_group.addButton(presence_no, 0)
        presence_group.addButton(presence_partial, 1)
        presence_group.addButton(presence_yes, 2)
        presence_group.addButton(presence_omit, 3)
        presence_row.addStretch()
        layout.addLayout(presence_row)

        relevance_lbl = QLabel("Relevance:")
        relevance_lbl.setStyleSheet(label_style + " font-weight: 600; margin-top: 8px;")
        layout.addWidget(relevance_lbl)
        relevance_row = QHBoxLayout()
        relevance_group = QButtonGroup(self)
        relevance_low = QRadioButton("Low")
        relevance_mod = QRadioButton("Moderate")
        relevance_high = QRadioButton("High")
        for rb in [relevance_low, relevance_mod, relevance_high]:
            rb.setStyleSheet(radio_style)
            relevance_row.addWidget(rb)
        relevance_group.addButton(relevance_low, 0)
        relevance_group.addButton(relevance_mod, 1)
        relevance_group.addButton(relevance_high, 2)
        relevance_row.addStretch()
        layout.addLayout(relevance_row)

        text_widgets = []
        subsection_data = []

        # === SUBSECTIONS COLLAPSIBLE WRAPPER ===
        if CollapsibleSection:
            subsections_section = CollapsibleSection("Subsections", start_collapsed=False)
            subsections_section.set_content_height(300)
            subsections_section._min_height = 100
            subsections_section._max_height = 600
            subsections_section.set_header_style("""
                QFrame {
                    background: rgba(5, 150, 105, 0.12);
                    border: 1px solid rgba(5, 150, 105, 0.3);
                    border-radius: 6px 6px 0 0;
                }
            """)
            subsections_section.set_title_style("""
                QLabel {
                    font-size: 18px;
                    font-weight: 600;
                    color: #059669;
                    background: transparent;
                    border: none;
                }
            """)
            subsections_content = QWidget()
            subsections_content.setStyleSheet("""
                QWidget {
                    background: rgba(5, 150, 105, 0.05);
                    border: 1px solid rgba(5, 150, 105, 0.2);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)
            subsections_content_layout = QVBoxLayout(subsections_content)
            subsections_content_layout.setContentsMargins(12, 10, 12, 10)
            subsections_content_layout.setSpacing(8)
            subsections_target_layout = subsections_content_layout
        else:
            subsections_target_layout = layout

        categories = [
            ("Anticipated Stressors", "rgba(254, 226, 226, 0.5)", "#991b1b", "r5_stress", [
                ("str_discharge", "Discharge/transition stress"),
                ("str_housing", "Housing uncertainty"),
                ("str_relationship", "Relationship strain"),
                ("str_financial", "Financial problems"),
                ("str_reduced_support", "Reduced support planned"),
            ]),
            ("Historical Pattern Under Stress", "rgba(254, 243, 199, 0.5)", "#b45309", "r5_pattern", [
                ("pat_deteriorates", "Deteriorates under stress"),
                ("pat_struggles_transitions", "Struggles during transitions"),
                ("pat_stress_incidents", "Stress has preceded incidents"),
            ]),
            ("Coping Capacity", "rgba(219, 234, 254, 0.5)", "#1d4ed8", "r5_coping", [
                ("cop_limited", "Limited coping skills"),
                ("cop_requires_containment", "Requires external containment"),
                ("cop_maladaptive", "Uses maladaptive coping (anger, avoidance)"),
                ("cop_effective", "Effective coping strategies (protective)"),
            ]),
            ("Substance Use as Coping", "rgba(243, 232, 255, 0.5)", "#7e22ce", "r5_substance", [
                ("sub_likely", "Substance use likely under stress"),
                ("sub_relapse_risk", "High relapse risk"),
                ("sub_history", "History of stress-linked substance use"),
            ]),
            ("Protective Factors", "rgba(220, 252, 231, 0.5)", "#15803d", "r5_protect", [
                ("prot_coping_demonstrated", "Demonstrated ability to cope"),
                ("prot_help_seeking", "Seeks help early"),
                ("prot_crisis_plan", "Rehearsed crisis plan"),
                ("prot_stable_supports", "Stable supports available"),
            ]),
        ]

        for cat_name, bg_color, text_color, cat_key, items in categories:
            cont = QFrame()
            cont.setStyleSheet(f"QFrame {{ background: {bg_color}; border: 2px solid {text_color}40; border-radius: 10px; }}")
            cont_layout = QVBoxLayout(cont)
            cont_layout.setContentsMargins(12, 10, 12, 10)
            cont_layout.setSpacing(8)

            lbl = QLabel(cat_name)
            lbl.setStyleSheet(f"font-size: 16px; font-weight: 600; color: {text_color}; background: transparent; border: none;")
            cont_layout.addWidget(lbl)

            txt = QTextEdit()
            txt.setPlaceholderText(f"Describe {cat_name.lower()}...")
            txt.setStyleSheet(input_style + " background: white;")
            txt.setMinimumHeight(58)
            cont_layout.addWidget(txt)
            cont_layout.addWidget(DragResizeBar(txt))
            text_widgets.append(txt)
            subsection_data.append((cat_name, txt))
            setattr(self, f"popup_{cat_key}", txt)

            if CollapsibleSection:
                section = CollapsibleSection("Select Applicable Indicators", start_collapsed=False)
                section.set_content_height(120)
                content = QWidget()
                content.setStyleSheet("QWidget { background: rgba(255,255,255,0.9); border-radius: 0 0 8px 8px; } QCheckBox { background: transparent; border: none; font-size: 13px; padding: 2px; }")
                cb_layout = QVBoxLayout(content)
                cb_layout.setContentsMargins(10, 6, 10, 6)
                cb_layout.setSpacing(3)

                checkboxes = []
                for item_key, item_label in items:
                    cb = QCheckBox(item_label)
                    cb.setProperty("r5_key", item_key)
                    cb.setProperty("r5_category", cat_key)
                    cb.stateChanged.connect(self._update_r5_stress_narrative)
                    cb_layout.addWidget(cb)
                    checkboxes.append(cb)

                section.set_content(content)
                cont_layout.addWidget(section)
                setattr(self, f"popup_{cat_key}_checkboxes", checkboxes)

            subsections_target_layout.addWidget(cont)

        if CollapsibleSection:
            subsections_section.set_content(subsections_content)
            layout.addWidget(subsections_section)

        layout.addStretch()
        setattr(self, f"popup_{key}_presence_group", presence_group)
        setattr(self, f"popup_{key}_relevance_group", relevance_group)
        setattr(self, f"popup_{key}_subsection_data", subsection_data)

        def generate():
            parts = [f"ITEM {code.upper()}: {description}"]
            for label, widget in subsection_data:
                text = widget.toPlainText().strip()
                if text:
                    parts.append(f"{label}: {text}")
            return "\n".join(parts)
        setattr(self, f"popup_{key}_generate", generate)

    def _update_r5_stress_narrative(self):
        """Update R5 text fields with subsection-aware, gender-sensitive narrative."""
        gender = getattr(self, '_current_gender', 'male')
        if gender == 'male':
            subj, obj, poss = 'he', 'him', 'his'
        else:
            subj, obj, poss = 'she', 'her', 'her'
        Subj = subj.capitalize()

        def join_items(items):
            if len(items) == 1:
                return items[0]
            elif len(items) == 2:
                return f"{items[0]} and {items[1]}"
            else:
                return ", ".join(items[:-1]) + f", and {items[-1]}"

        # ============ 1. ANTICIPATED STRESSORS (all vulnerability) ============
        str_items = {"str_discharge": "discharge/transition stress", "str_housing": "housing uncertainty", "str_relationship": "relationship strain", "str_financial": "financial problems", "str_reduced_support": "reduced support planned"}

        str_cbs = getattr(self, "popup_r5_stress_checkboxes", [])
        str_text = getattr(self, "popup_r5_stress", None)

        str_checked = []
        for cb in str_cbs:
            if cb.isChecked():
                key = cb.property("r5_key")
                if key in str_items:
                    str_checked.append(str_items[key])

        if str_text:
            if str_checked:
                str_text.setPlainText(f"Anticipated stressors include {join_items(str_checked)}.")
            else:
                str_text.clear()

        # ============ 2. HISTORICAL PATTERN UNDER STRESS (all vulnerability) ============
        pat_items = {"pat_deteriorates": "deteriorating under stress", "pat_struggles_transitions": "struggling during transitions", "pat_stress_incidents": "stress preceding incidents"}

        pat_cbs = getattr(self, "popup_r5_pattern_checkboxes", [])
        pat_text = getattr(self, "popup_r5_pattern", None)

        pat_checked = []
        for cb in pat_cbs:
            if cb.isChecked():
                key = cb.property("r5_key")
                if key in pat_items:
                    pat_checked.append(pat_items[key])

        if pat_text:
            if pat_checked:
                pat_text.setPlainText(f"Historical patterns include {join_items(pat_checked)}.")
            else:
                pat_text.clear()

        # ============ 3. COPING CAPACITY (mixed) ============
        cop_vuln = {"cop_limited": "limited coping skills", "cop_requires_containment": "requiring external containment", "cop_maladaptive": "using maladaptive coping strategies"}
        cop_prot = {"cop_effective": "effective coping strategies"}

        cop_cbs = getattr(self, "popup_r5_coping_checkboxes", [])
        cop_text = getattr(self, "popup_r5_coping", None)

        cop_v, cop_p = [], []
        for cb in cop_cbs:
            if cb.isChecked():
                key = cb.property("r5_key")
                if key in cop_vuln:
                    cop_v.append(cop_vuln[key])
                elif key in cop_prot:
                    cop_p.append(cop_prot[key])

        if cop_text:
            sentences = []
            if cop_v:
                sentences.append(f"Coping capacity concerns include {join_items(cop_v)}.")
            if cop_p:
                if sentences:
                    sentences.append(f"However, {subj} has {join_items(cop_p)}.")
                else:
                    sentences.append(f"{Subj} has {join_items(cop_p)}.")
            cop_text.setPlainText(" ".join(sentences)) if sentences else cop_text.clear()

        # ============ 4. SUBSTANCE USE AS COPING (all vulnerability) ============
        sub_items = {"sub_likely": "substance use likely under stress", "sub_relapse_risk": "high relapse risk", "sub_history": "a history of stress-linked substance use"}

        sub_cbs = getattr(self, "popup_r5_substance_checkboxes", [])
        sub_text = getattr(self, "popup_r5_substance", None)

        sub_checked = []
        for cb in sub_cbs:
            if cb.isChecked():
                key = cb.property("r5_key")
                if key in sub_items:
                    sub_checked.append(sub_items[key])

        if sub_text:
            if sub_checked:
                sub_text.setPlainText(f"Substance-related coping concerns include {join_items(sub_checked)}.")
            else:
                sub_text.clear()

        # ============ 5. PROTECTIVE FACTORS (all protective) ============
        prot_items = {"prot_coping_demonstrated": "demonstrated ability to cope", "prot_help_seeking": "seeking help early", "prot_crisis_plan": "a rehearsed crisis plan", "prot_stable_supports": "stable supports available"}

        prot_cbs = getattr(self, "popup_r5_protect_checkboxes", [])
        prot_text = getattr(self, "popup_r5_protect", None)

        prot_checked = []
        for cb in prot_cbs:
            if cb.isChecked():
                key = cb.property("r5_key")
                if key in prot_items:
                    prot_checked.append(prot_items[key])

        if prot_text:
            if prot_checked:
                prot_text.setPlainText(f"Protective factors include {join_items(prot_checked)}.")
            else:
                prot_text.clear()

    def _build_popup_formulation(self):
        """Build violence risk formulation popup."""
        container, layout = self._create_popup_container("formulation")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("Violence Risk Formulation:")
        lbl.setStyleSheet("font-size: 18px; font-weight: 600; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("Provide a narrative formulation integrating the HCR-20 factors, explaining the individual's pathway to violence and key risk drivers.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        self.popup_formulation_text = QTextEdit()
        self.popup_formulation_text.setPlaceholderText("Describe the individual's history, developmental factors, mental health, substance use, relationships, and how these contribute to violence risk...")
        self.popup_formulation_text.setStyleSheet(input_style)
        self.popup_formulation_text.setMinimumHeight(216)  # Reduced by 10% (was 240)
        layout.addWidget(self.popup_formulation_text)

        layout.addStretch()

        def generate():
            return self.popup_formulation_text.toPlainText().strip()

        self._connect_preview_updates("formulation", [self.popup_formulation_text])
        self._add_send_button(layout, "formulation", generate)

    # ================================================================
    #  RISK FORMULATION SECTION - 11 Domains with Auto-Population
    # ================================================================

    def _build_popup_nature_of_risk(self):
        """Build Nature of Risk popup with harm type checkboxes."""
        container, layout = self._create_popup_container("scenario_nature")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("1. Nature of Risk")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("Specify what kind of harm is likely, not just 'violence'. Select all applicable types.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Harm type container
        harm_container = QFrame()
        harm_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        harm_layout = QVBoxLayout(harm_container)
        harm_layout.setContentsMargins(12, 10, 12, 10)
        harm_layout.setSpacing(8)

        harm_lbl = QLabel("Type of Harm")
        harm_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        harm_layout.addWidget(harm_lbl)

        HARM_TYPES = [
            ("harm_physical_general", "Physical violence (general)"),
            ("harm_physical_targeted", "Physical violence (targeted at specific person)"),
            ("harm_domestic", "Domestic violence"),
            ("harm_threatening", "Threatening or intimidating behaviour"),
            ("harm_weapon", "Weapon-related risk"),
            ("harm_sexual", "Sexual violence"),
            ("harm_arson", "Arson / property damage"),
            ("harm_institutional", "Institutional aggression (staff, patients)"),
            ("harm_stalking", "Stalking / harassment"),
        ]

        self.popup_nature_checkboxes = []
        for item_key, item_label in HARM_TYPES:
            cb = QCheckBox(item_label)
            cb.setProperty("nature_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_nature_narrative)
            harm_layout.addWidget(cb)
            self.popup_nature_checkboxes.append(cb)

        layout.addWidget(harm_container)

        # Victim container
        victim_container = QFrame()
        victim_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        victim_layout = QVBoxLayout(victim_container)
        victim_layout.setContentsMargins(12, 10, 12, 10)
        victim_layout.setSpacing(8)

        victim_lbl = QLabel("Potential Victims")
        victim_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        victim_layout.addWidget(victim_lbl)

        VICTIM_TYPES = [
            ("victim_known", "Known others (family, partners, acquaintances)"),
            ("victim_strangers", "Strangers"),
            ("victim_staff", "Staff / professionals"),
            ("victim_patients", "Co-patients / service users"),
            ("victim_authority", "Authority figures"),
            ("victim_children", "Vulnerable groups (children, elderly)"),
        ]

        self.popup_victim_checkboxes = []
        for item_key, item_label in VICTIM_TYPES:
            cb = QCheckBox(item_label)
            cb.setProperty("victim_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_nature_narrative)
            victim_layout.addWidget(cb)
            self.popup_victim_checkboxes.append(cb)

        layout.addWidget(victim_container)

        # Motivation container
        motiv_container = QFrame()
        motiv_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        motiv_layout = QVBoxLayout(motiv_container)
        motiv_layout.setContentsMargins(12, 10, 12, 10)
        motiv_layout.setSpacing(8)

        motiv_lbl = QLabel("Likely Motivation")
        motiv_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        motiv_layout.addWidget(motiv_lbl)

        MOTIVATIONS = [
            ("motiv_impulsive", "Impulsive / reactive (emotional dysregulation)"),
            ("motiv_instrumental", "Instrumental / goal-directed"),
            ("motiv_paranoid", "Paranoid / persecutory beliefs"),
            ("motiv_command", "Response to command hallucinations"),
            ("motiv_grievance", "Grievance / revenge"),
            ("motiv_territorial", "Territorial / defensive"),
            ("motiv_substance", "Substance-related disinhibition"),
        ]

        self.popup_motivation_checkboxes = []
        for item_key, item_label in MOTIVATIONS:
            cb = QCheckBox(item_label)
            cb.setProperty("motiv_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_nature_narrative)
            motiv_layout.addWidget(cb)
            self.popup_motivation_checkboxes.append(cb)

        layout.addWidget(motiv_container)

        layout.addStretch()

        def generate():
            return self._generate_nature_narrative()

        self._connect_preview_updates("scenario_nature", [])
        self._add_send_button(layout, "scenario_nature", generate)

    def _generate_nature_narrative(self):
        """Generate Nature of Risk narrative from checkboxes."""
        harm_phrases = {
            "harm_physical_general": "general physical aggression",
            "harm_physical_targeted": "targeted physical violence towards specific individuals",
            "harm_domestic": "domestic violence",
            "harm_threatening": "threatening or intimidating behaviour",
            "harm_weapon": "weapon-related violence",
            "harm_sexual": "sexual violence",
            "harm_arson": "fire-setting or property damage",
            "harm_institutional": "institutional aggression towards staff or patients",
            "harm_stalking": "stalking or harassment behaviour",
        }

        victim_phrases = {
            "victim_known": "known others (family, partners, acquaintances)",
            "victim_strangers": "strangers",
            "victim_staff": "staff and professionals",
            "victim_patients": "co-patients or service users",
            "victim_authority": "authority figures",
            "victim_children": "vulnerable groups including children or elderly",
        }

        motiv_phrases = {
            "motiv_impulsive": "impulsive physical aggression during periods of emotional dysregulation",
            "motiv_instrumental": "instrumental or goal-directed violence",
            "motiv_paranoid": "violence driven by paranoid or persecutory beliefs",
            "motiv_command": "acting on command hallucinations",
            "motiv_grievance": "grievance-based or retaliatory violence",
            "motiv_territorial": "territorial or defensive aggression",
            "motiv_substance": "violence during substance-related disinhibition",
        }

        harm_selected = []
        for cb in getattr(self, 'popup_nature_checkboxes', []):
            if cb.isChecked():
                key = cb.property("nature_key")
                if key in harm_phrases:
                    harm_selected.append(harm_phrases[key])

        victim_selected = []
        for cb in getattr(self, 'popup_victim_checkboxes', []):
            if cb.isChecked():
                key = cb.property("victim_key")
                if key in victim_phrases:
                    victim_selected.append(victim_phrases[key])

        motiv_selected = []
        for cb in getattr(self, 'popup_motivation_checkboxes', []):
            if cb.isChecked():
                key = cb.property("motiv_key")
                if key in motiv_phrases:
                    motiv_selected.append(motiv_phrases[key])

        parts = []
        if harm_selected:
            harm_str = ", ".join(harm_selected[:-1]) + " and " + harm_selected[-1] if len(harm_selected) > 1 else harm_selected[0]
            parts.append(f"Risk relates primarily to {harm_str}")

        if victim_selected:
            victim_str = ", ".join(victim_selected[:-1]) + " and " + victim_selected[-1] if len(victim_selected) > 1 else victim_selected[0]
            parts.append(f"towards {victim_str}")

        if motiv_selected:
            motiv_str = motiv_selected[0]  # Primary motivation
            if len(parts) > 0:
                parts.append(f"characterised by {motiv_str}")
            else:
                parts.append(f"Risk is characterised by {motiv_str}")

        return " ".join(parts) + "." if parts else ""

    def _update_nature_narrative(self):
        """Update Nature of Risk card from checkboxes."""
        self._update_preview("scenario_nature")

    def _build_popup_severity(self):
        """Build Severity popup with harm severity anchors."""
        container, layout = self._create_popup_container("scenario_severity")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 15px; padding: 4px; }"

        lbl = QLabel("2. Severity")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("How serious could the harm be? Provide realistic upper bounds, not worst-case fantasy.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Severity level container
        severity_container = QFrame()
        severity_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        severity_layout = QVBoxLayout(severity_container)
        severity_layout.setContentsMargins(12, 10, 12, 10)
        severity_layout.setSpacing(8)

        severity_lbl = QLabel("Severity Level")
        severity_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        severity_layout.addWidget(severity_lbl)

        self.popup_severity_group = QButtonGroup(self)

        severity_low = QRadioButton("Low — minor injury, verbal threats, transient aggression")
        severity_low.setStyleSheet(radio_style)
        self.popup_severity_group.addButton(severity_low, 0)
        severity_layout.addWidget(severity_low)

        severity_mod = QRadioButton("Moderate — assault causing injury, repeated threats, use of force")
        severity_mod.setStyleSheet(radio_style)
        self.popup_severity_group.addButton(severity_mod, 1)
        severity_layout.addWidget(severity_mod)

        severity_high = QRadioButton("High — serious injury, weapon use, high vulnerability victims")
        severity_high.setStyleSheet(radio_style)
        self.popup_severity_group.addButton(severity_high, 2)
        severity_layout.addWidget(severity_high)

        self.popup_severity_group.buttonClicked.connect(self._update_severity_narrative)
        layout.addWidget(severity_container)

        # Additional factors container
        factors_container = QFrame()
        factors_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        factors_layout = QVBoxLayout(factors_container)
        factors_layout.setContentsMargins(12, 10, 12, 10)
        factors_layout.setSpacing(8)

        factors_lbl = QLabel("Severity Factors")
        factors_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        factors_layout.addWidget(factors_lbl)

        SEVERITY_FACTORS = [
            ("sev_hist_serious", "History of assaults involving significant force"),
            ("sev_limited_inhibition", "Limited inhibition when unwell"),
            ("sev_weapon_history", "Previous weapon use"),
            ("sev_vulnerable_victims", "Risk to vulnerable victims"),
            ("sev_escalation_pattern", "Pattern of escalation in violence"),
            ("sev_lack_remorse", "Lack of remorse following violence"),
        ]

        self.popup_severity_factor_checkboxes = []
        for item_key, item_label in SEVERITY_FACTORS:
            cb = QCheckBox(item_label)
            cb.setProperty("sev_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_severity_narrative)
            factors_layout.addWidget(cb)
            self.popup_severity_factor_checkboxes.append(cb)

        layout.addWidget(factors_container)

        layout.addStretch()

        def generate():
            return self._generate_severity_narrative()

        self._connect_preview_updates("scenario_severity", [])
        self._add_send_button(layout, "scenario_severity", generate)

    def _generate_severity_narrative(self):
        """Generate Severity narrative from selections."""
        severity_level = ""
        btn = self.popup_severity_group.checkedButton()
        if btn:
            btn_id = self.popup_severity_group.id(btn)
            if btn_id == 0:
                severity_level = "low"
            elif btn_id == 1:
                severity_level = "moderate"
            elif btn_id == 2:
                severity_level = "high"

        factor_phrases = {
            "sev_hist_serious": "history of assaults involving significant force",
            "sev_limited_inhibition": "limited inhibition when unwell",
            "sev_weapon_history": "previous weapon use",
            "sev_vulnerable_victims": "risk to vulnerable victims",
            "sev_escalation_pattern": "pattern of escalation in violence",
            "sev_lack_remorse": "lack of remorse following violence",
        }

        factors_selected = []
        for cb in getattr(self, 'popup_severity_factor_checkboxes', []):
            if cb.isChecked():
                key = cb.property("sev_key")
                if key in factor_phrases:
                    factors_selected.append(factor_phrases[key])

        parts = []
        if severity_level:
            if severity_level == "low":
                parts.append("Potential severity is assessed as low, with risk primarily involving minor injury, verbal threats, or transient aggression")
            elif severity_level == "moderate":
                parts.append("Potential severity is assessed as moderate, with risk of assault causing injury, repeated threats, or use of force")
            elif severity_level == "high":
                parts.append("Potential severity is assessed as high, with risk of serious injury, weapon use, or harm to high vulnerability victims")

        if factors_selected:
            factors_str = ", ".join(factors_selected[:-1]) + " and " + factors_selected[-1] if len(factors_selected) > 1 else factors_selected[0]
            if parts:
                parts.append(f"given {factors_str}")
            else:
                parts.append(f"Severity factors include {factors_str}")

        return " ".join(parts) + "." if parts else ""

    def _update_severity_narrative(self):
        """Update Severity card from selections."""
        self._update_preview("scenario_severity")

    def _build_popup_imminence(self):
        """Build Imminence popup with time-sensitive triggers."""
        container, layout = self._create_popup_container("scenario_imminence")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("3. Imminence")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("How soon could risk escalate? Consider pending transitions and changes to protective factors.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Trigger presence container
        trigger_container = QFrame()
        trigger_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        trigger_layout = QVBoxLayout(trigger_container)
        trigger_layout.setContentsMargins(12, 10, 12, 10)
        trigger_layout.setSpacing(8)

        trigger_lbl = QLabel("Current Trigger Status")
        trigger_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        trigger_layout.addWidget(trigger_lbl)

        TRIGGER_STATUS = [
            ("trig_present", "Triggers already present"),
            ("trig_emerging", "Triggers emerging / building"),
            ("trig_absent", "No current triggers identified"),
        ]

        self.popup_trigger_group = QButtonGroup(self)
        self.popup_trigger_radios = {}
        for idx, (item_key, item_label) in enumerate(TRIGGER_STATUS):
            rb = QRadioButton(item_label)
            rb.setProperty("trig_key", item_key)
            rb.setStyleSheet("QRadioButton { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            rb.toggled.connect(self._update_imminence_narrative)
            trigger_layout.addWidget(rb)
            self.popup_trigger_group.addButton(rb, idx)
            self.popup_trigger_radios[item_key] = rb

        layout.addWidget(trigger_container)

        # Transition container
        transition_container = QFrame()
        transition_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        transition_layout = QVBoxLayout(transition_container)
        transition_layout.setContentsMargins(12, 10, 12, 10)
        transition_layout.setSpacing(8)

        transition_lbl = QLabel("Pending Transitions")
        transition_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        transition_layout.addWidget(transition_lbl)

        TRANSITIONS = [
            ("trans_discharge", "Discharge pending"),
            ("trans_leave", "Leave / unescorted access pending"),
            ("trans_reduced_supervision", "Supervision reduction pending"),
            ("trans_accommodation", "Accommodation change pending"),
            ("trans_relationship", "Relationship change anticipated"),
            ("trans_legal", "Legal proceedings pending"),
        ]

        self.popup_transition_checkboxes = []
        for item_key, item_label in TRANSITIONS:
            cb = QCheckBox(item_label)
            cb.setProperty("trans_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_imminence_narrative)
            transition_layout.addWidget(cb)
            self.popup_transition_checkboxes.append(cb)

        layout.addWidget(transition_container)

        # Protective factor changes
        protect_container = QFrame()
        protect_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(34, 197, 94, 0.4);
                border-radius: 10px;
            }
        """)
        protect_layout = QVBoxLayout(protect_container)
        protect_layout.setContentsMargins(12, 10, 12, 10)
        protect_layout.setSpacing(8)

        protect_lbl = QLabel("Protective Factor Changes")
        protect_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #166534; background: transparent; border: none;")
        protect_layout.addWidget(protect_lbl)

        PROTECT_CHANGES = [
            ("prot_reducing", "Protective factors about to reduce"),
            ("prot_stable", "Protective factors stable"),
            ("prot_increasing", "Protective factors increasing"),
        ]

        self.popup_protect_group = QButtonGroup(self)
        self.popup_protect_radios = {}
        for idx, (item_key, item_label) in enumerate(PROTECT_CHANGES):
            rb = QRadioButton(item_label)
            rb.setProperty("prot_key", item_key)
            rb.setStyleSheet("QRadioButton { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            rb.toggled.connect(self._update_imminence_narrative)
            protect_layout.addWidget(rb)
            self.popup_protect_group.addButton(rb, idx)
            self.popup_protect_radios[item_key] = rb

        layout.addWidget(protect_container)

        layout.addStretch()

        def generate():
            return self._generate_imminence_narrative()

        self._connect_preview_updates("scenario_imminence", [])
        self._add_send_button(layout, "scenario_imminence", generate)

    def _generate_imminence_narrative(self):
        """Generate Imminence narrative from selections."""
        parts = []

        # Trigger status (radio buttons)
        trigger_radios = getattr(self, 'popup_trigger_radios', {})
        for key, rb in trigger_radios.items():
            if rb.isChecked():
                if key == "trig_present":
                    parts.append("Known risk triggers are currently present")
                elif key == "trig_emerging":
                    parts.append("Risk triggers appear to be emerging")
                elif key == "trig_absent":
                    parts.append("No specific triggers are currently identified")
                break  # Only one can be selected

        # Transitions (checkboxes - multiple can be selected)
        transitions = []
        trans_phrases = {
            "trans_discharge": "the immediate post-discharge period",
            "trans_leave": "progression to unescorted leave",
            "trans_reduced_supervision": "reduction in supervision",
            "trans_accommodation": "change in accommodation",
            "trans_relationship": "anticipated relationship changes",
            "trans_legal": "pending legal proceedings",
        }
        for cb in getattr(self, 'popup_transition_checkboxes', []):
            if cb.isChecked():
                key = cb.property("trans_key")
                if key in trans_phrases:
                    transitions.append(trans_phrases[key])

        if transitions:
            trans_str = ", ".join(transitions[:-1]) + " and " + transitions[-1] if len(transitions) > 1 else transitions[0]
            parts.append(f"Risk may become imminent during {trans_str} when supervision reduces and stressors increase")

        # Protective factor changes (radio buttons)
        protect_radios = getattr(self, 'popup_protect_radios', {})
        for key, rb in protect_radios.items():
            if rb.isChecked():
                if key == "prot_reducing":
                    parts.append("Protective factors are about to reduce, potentially increasing imminence")
                elif key == "prot_stable":
                    parts.append("Protective factors remain stable at this time")
                elif key == "prot_increasing":
                    parts.append("Protective factors are increasing, reducing immediate risk")
                break  # Only one can be selected

        return ". ".join(parts) + "." if parts else ""

    def _update_imminence_narrative(self):
        """Update Imminence card from selections."""
        self._update_preview("scenario_imminence")

    def _build_popup_frequency(self):
        """Build Frequency popup with pattern anchors."""
        container, layout = self._create_popup_container("scenario_frequency")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 15px; padding: 4px; }"

        lbl = QLabel("4. Frequency")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("How often might risk events occur? Focus on patterns, not counts.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Pattern container
        pattern_container = QFrame()
        pattern_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        pattern_layout = QVBoxLayout(pattern_container)
        pattern_layout.setContentsMargins(12, 10, 12, 10)
        pattern_layout.setSpacing(8)

        pattern_lbl = QLabel("Risk Pattern")
        pattern_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        pattern_layout.addWidget(pattern_lbl)

        self.popup_frequency_group = QButtonGroup(self)

        freq_episodic = QRadioButton("Episodic — linked to specific triggers")
        freq_episodic.setStyleSheet(radio_style)
        self.popup_frequency_group.addButton(freq_episodic, 0)
        pattern_layout.addWidget(freq_episodic)

        freq_clustered = QRadioButton("Clustered — during relapse / stress periods")
        freq_clustered.setStyleSheet(radio_style)
        self.popup_frequency_group.addButton(freq_clustered, 1)
        pattern_layout.addWidget(freq_clustered)

        freq_persistent = QRadioButton("Persistent — background risk")
        freq_persistent.setStyleSheet(radio_style)
        self.popup_frequency_group.addButton(freq_persistent, 2)
        pattern_layout.addWidget(freq_persistent)

        freq_rare = QRadioButton("Rare but severe")
        freq_rare.setStyleSheet(radio_style)
        self.popup_frequency_group.addButton(freq_rare, 3)
        pattern_layout.addWidget(freq_rare)

        self.popup_frequency_group.buttonClicked.connect(self._update_frequency_narrative)
        layout.addWidget(pattern_container)

        # Context container
        context_container = QFrame()
        context_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        context_layout = QVBoxLayout(context_container)
        context_layout.setContentsMargins(12, 10, 12, 10)
        context_layout.setSpacing(8)

        context_lbl = QLabel("Trigger Contexts")
        context_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        context_layout.addWidget(context_lbl)

        CONTEXTS = [
            ("ctx_stress", "Acute stress periods"),
            ("ctx_relapse", "Mental health relapse"),
            ("ctx_substance", "Substance intoxication"),
            ("ctx_interpersonal", "Interpersonal conflict"),
            ("ctx_frustration", "Frustration / goal blocking"),
            ("ctx_perceived_threat", "Perceived threat or provocation"),
        ]

        self.popup_context_checkboxes = []
        for item_key, item_label in CONTEXTS:
            cb = QCheckBox(item_label)
            cb.setProperty("ctx_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_frequency_narrative)
            context_layout.addWidget(cb)
            self.popup_context_checkboxes.append(cb)

        layout.addWidget(context_container)

        layout.addStretch()

        def generate():
            return self._generate_frequency_narrative()

        self._connect_preview_updates("scenario_frequency", [])
        self._add_send_button(layout, "scenario_frequency", generate)

    def _generate_frequency_narrative(self):
        """Generate Frequency narrative from selections."""
        sentences = []

        # Risk pattern - discrete sentence
        btn = self.popup_frequency_group.checkedButton()
        if btn:
            btn_id = self.popup_frequency_group.id(btn)
            if btn_id == 0:
                sentences.append("Risk has historically occurred in episodic patterns linked to specific triggers.")
            elif btn_id == 1:
                sentences.append("Risk has historically occurred in episodic clusters during periods of acute stress or relapse.")
            elif btn_id == 2:
                sentences.append("Risk presents as a persistent background behaviour rather than discrete episodes.")
            elif btn_id == 3:
                sentences.append("Risk events are rare but potentially severe when they occur.")

        # Trigger contexts - separate sentence
        ctx_phrases = {
            "ctx_stress": "periods of acute stress",
            "ctx_relapse": "mental health relapse",
            "ctx_substance": "substance intoxication",
            "ctx_interpersonal": "interpersonal conflict",
            "ctx_frustration": "frustration or goal blocking",
            "ctx_perceived_threat": "perceived threat or provocation",
        }

        contexts = []
        for cb in getattr(self, 'popup_context_checkboxes', []):
            if cb.isChecked():
                key = cb.property("ctx_key")
                if key in ctx_phrases:
                    contexts.append(ctx_phrases[key])

        if contexts:
            ctx_str = ", ".join(contexts[:-1]) + " and " + contexts[-1] if len(contexts) > 1 else contexts[0]
            sentences.append(f"The context of these triggers include {ctx_str}.")

        return " ".join(sentences)

    def _update_frequency_narrative(self):
        """Update Frequency card from selections."""
        self._update_preview("scenario_frequency")

    def _build_popup_likelihood(self):
        """Build Likelihood popup with conditional probability anchors."""
        container, layout = self._create_popup_container("scenario_likelihood")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 15px; padding: 4px; }"

        lbl = QLabel("5. Likelihood")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("How plausible is the scenario? This is conditional likelihood, not absolute probability.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Baseline likelihood container
        baseline_container = QFrame()
        baseline_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        baseline_layout = QVBoxLayout(baseline_container)
        baseline_layout.setContentsMargins(12, 10, 12, 10)
        baseline_layout.setSpacing(8)

        baseline_lbl = QLabel("Baseline Likelihood")
        baseline_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        baseline_layout.addWidget(baseline_lbl)

        self.popup_likelihood_group = QButtonGroup(self)

        like_low = QRadioButton("Low — unless specific triggers occur")
        like_low.setStyleSheet(radio_style)
        self.popup_likelihood_group.addButton(like_low, 0)
        baseline_layout.addWidget(like_low)

        like_mod = QRadioButton("Moderate — if current stability maintained")
        like_mod.setStyleSheet(radio_style)
        self.popup_likelihood_group.addButton(like_mod, 1)
        baseline_layout.addWidget(like_mod)

        like_high = QRadioButton("High — when known risk pattern re-emerges")
        like_high.setStyleSheet(radio_style)
        self.popup_likelihood_group.addButton(like_high, 2)
        baseline_layout.addWidget(like_high)

        self.popup_likelihood_group.buttonClicked.connect(self._update_likelihood_narrative)
        layout.addWidget(baseline_container)

        # Conditional factors container
        cond_container = QFrame()
        cond_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        cond_layout = QVBoxLayout(cond_container)
        cond_layout.setContentsMargins(12, 10, 12, 10)
        cond_layout.setSpacing(8)

        cond_lbl = QLabel("Likelihood Increases If...")
        cond_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        cond_layout.addWidget(cond_lbl)

        CONDITIONS = [
            ("cond_med_lapse", "Medication adherence lapses"),
            ("cond_conflict", "Interpersonal conflict escalates"),
            ("cond_substance", "Substance use resumes"),
            ("cond_supervision_reduces", "Supervision reduces"),
            ("cond_symptoms_return", "Symptoms of mental disorder return"),
            ("cond_support_loss", "Support network weakens"),
            ("cond_stress_increases", "Life stressors increase"),
        ]

        self.popup_likelihood_cond_checkboxes = []
        for item_key, item_label in CONDITIONS:
            cb = QCheckBox(item_label)
            cb.setProperty("cond_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_likelihood_narrative)
            cond_layout.addWidget(cb)
            self.popup_likelihood_cond_checkboxes.append(cb)

        layout.addWidget(cond_container)

        layout.addStretch()

        def generate():
            return self._generate_likelihood_narrative()

        self._connect_preview_updates("scenario_likelihood", [])
        self._add_send_button(layout, "scenario_likelihood", generate)

    def _generate_likelihood_narrative(self):
        """Generate Likelihood narrative from selections."""
        parts = []

        btn = self.popup_likelihood_group.checkedButton()
        if btn:
            btn_id = self.popup_likelihood_group.id(btn)
            if btn_id == 0:
                parts.append("Likelihood is assessed as low unless specific triggers occur")
            elif btn_id == 1:
                parts.append("Likelihood is assessed as moderate if current stability is maintained")
            elif btn_id == 2:
                parts.append("Likelihood is assessed as high when the known risk pattern re-emerges")

        cond_phrases = {
            "cond_med_lapse": "medication adherence lapses",
            "cond_conflict": "interpersonal conflict escalates",
            "cond_substance": "substance use resumes",
            "cond_supervision_reduces": "supervision reduces",
            "cond_symptoms_return": "symptoms of mental disorder return",
            "cond_support_loss": "support network weakens",
            "cond_stress_increases": "life stressors increase",
        }

        conditions = []
        for cb in getattr(self, 'popup_likelihood_cond_checkboxes', []):
            if cb.isChecked():
                key = cb.property("cond_key")
                if key in cond_phrases:
                    conditions.append(cond_phrases[key])

        if conditions:
            cond_str = ", ".join(conditions[:-1]) + " or " + conditions[-1] if len(conditions) > 1 else conditions[0]
            parts.append(f"but increases significantly if {cond_str}")

        narrative = ", ".join(parts) + "." if parts else ""
        return narrative.replace(",.", ".")

    def _update_likelihood_narrative(self):
        """Update Likelihood card from selections."""
        self._update_preview("scenario_likelihood")

    def _build_popup_risk_enhancing(self):
        """Build Risk-Enhancing Factors popup with auto-population from C and R items."""
        container, layout = self._create_popup_container("risk_enhancing")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("6. Risk-Enhancing Factors")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("What makes things worse? These are pulled from Clinical and Risk Management items.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Clinical enhancers container
        clinical_container = QFrame()
        clinical_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        clinical_layout = QVBoxLayout(clinical_container)
        clinical_layout.setContentsMargins(12, 10, 12, 10)
        clinical_layout.setSpacing(8)

        clinical_lbl = QLabel("Clinical Factors (from C1-C5)")
        clinical_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        clinical_layout.addWidget(clinical_lbl)

        CLINICAL_ENHANCERS = [
            ("enh_poor_insight", "Poor insight into illness or risk"),
            ("enh_violent_ideation", "Active violent ideation or intent"),
            ("enh_active_symptoms", "Active symptoms (psychosis, mania)"),
            ("enh_instability", "Affective or behavioural instability"),
            ("enh_poor_treatment", "Poor treatment or supervision response"),
        ]

        self.popup_clinical_enhancer_checkboxes = []
        for item_key, item_label in CLINICAL_ENHANCERS:
            cb = QCheckBox(item_label)
            cb.setProperty("enh_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_enhancing_narrative)
            clinical_layout.addWidget(cb)
            self.popup_clinical_enhancer_checkboxes.append(cb)

        layout.addWidget(clinical_container)

        # Situational enhancers container
        situational_container = QFrame()
        situational_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        situational_layout = QVBoxLayout(situational_container)
        situational_layout.setContentsMargins(12, 10, 12, 10)
        situational_layout.setSpacing(8)

        situational_lbl = QLabel("Situational Factors (from R1-R5)")
        situational_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        situational_layout.addWidget(situational_lbl)

        SITUATIONAL_ENHANCERS = [
            ("enh_poor_plan", "Inadequate professional services or plans"),
            ("enh_unstable_living", "Unstable or unsupportive living situation"),
            ("enh_poor_support", "Lack of personal support"),
            ("enh_non_compliance", "Non-compliance with treatment or supervision"),
            ("enh_poor_coping", "Poor stress tolerance or coping"),
        ]

        self.popup_situational_enhancer_checkboxes = []
        for item_key, item_label in SITUATIONAL_ENHANCERS:
            cb = QCheckBox(item_key)
            cb.setText(item_label)
            cb.setProperty("enh_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_enhancing_narrative)
            situational_layout.addWidget(cb)
            self.popup_situational_enhancer_checkboxes.append(cb)

        layout.addWidget(situational_container)

        # Other enhancers
        other_container = QFrame()
        other_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        other_layout = QVBoxLayout(other_container)
        other_layout.setContentsMargins(12, 10, 12, 10)
        other_layout.setSpacing(8)

        other_lbl = QLabel("Other Enhancing Factors")
        other_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        other_layout.addWidget(other_lbl)

        OTHER_ENHANCERS = [
            ("enh_substance", "Substance use"),
            ("enh_conflict", "Conflictual relationships"),
            ("enh_access_victims", "Access to victims or triggers"),
            ("enh_transitions", "Stressful transitions"),
            ("enh_loss_supervision", "Loss of supervision"),
        ]

        self.popup_other_enhancer_checkboxes = []
        for item_key, item_label in OTHER_ENHANCERS:
            cb = QCheckBox(item_label)
            cb.setProperty("enh_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_enhancing_narrative)
            other_layout.addWidget(cb)
            self.popup_other_enhancer_checkboxes.append(cb)

        layout.addWidget(other_container)

        layout.addStretch()

        def generate():
            return self._generate_enhancing_narrative()

        self._connect_preview_updates("risk_enhancing", [])
        self._add_send_button(layout, "risk_enhancing", generate)

    def _generate_enhancing_narrative(self):
        """Generate Risk-Enhancing Factors narrative from selections."""
        # Clinical factor phrases
        clinical_phrases = {
            "enh_poor_insight": "poor insight into illness or risk",
            "enh_violent_ideation": "active violent ideation or intent",
            "enh_active_symptoms": "active symptoms of mental disorder",
            "enh_instability": "affective or behavioural instability",
            "enh_poor_treatment": "poor treatment or supervision response",
        }

        # Situational factor phrases
        situational_phrases = {
            "enh_poor_plan": "inadequate professional services or plans",
            "enh_unstable_living": "unstable or unsupportive living situation",
            "enh_poor_support": "lack of personal support",
            "enh_non_compliance": "non-compliance with treatment or supervision",
            "enh_poor_coping": "poor stress tolerance or coping",
        }

        # Other/Personal factor phrases
        other_phrases = {
            "enh_substance": "ongoing substance use",
            "enh_conflict": "conflictual relationships",
            "enh_access_victims": "access to victims or known triggers",
            "enh_transitions": "stressful transitions",
            "enh_loss_supervision": "anticipated loss of supervision",
        }

        # Collect clinical factors
        clinical_factors = []
        for cb in getattr(self, 'popup_clinical_enhancer_checkboxes', []):
            if cb.isChecked():
                key = cb.property("enh_key")
                if key in clinical_phrases:
                    clinical_factors.append(clinical_phrases[key])

        # Collect situational factors
        situational_factors = []
        for cb in getattr(self, 'popup_situational_enhancer_checkboxes', []):
            if cb.isChecked():
                key = cb.property("enh_key")
                if key in situational_phrases:
                    situational_factors.append(situational_phrases[key])

        # Collect other factors
        other_factors = []
        for cb in getattr(self, 'popup_other_enhancer_checkboxes', []):
            if cb.isChecked():
                key = cb.property("enh_key")
                if key in other_phrases:
                    other_factors.append(other_phrases[key])

        sentences = []

        # Build clinical sentence
        if clinical_factors:
            clinical_str = ", ".join(clinical_factors[:-1]) + " and " + clinical_factors[-1] if len(clinical_factors) > 1 else clinical_factors[0]
            sentences.append(f"Risk is enhanced by clinical factors, including {clinical_str}.")

        # Build situational sentence
        if situational_factors:
            situational_str = ", ".join(situational_factors[:-1]) + " and " + situational_factors[-1] if len(situational_factors) > 1 else situational_factors[0]
            sentences.append(f"Situational factors also increase the risk such as {situational_str}.")

        # Build other factors sentence with conditional prefix
        if other_factors:
            other_str = ", ".join(other_factors[:-1]) + " and " + other_factors[-1] if len(other_factors) > 1 else other_factors[0]
            if clinical_factors and situational_factors:
                # Both previous sections have content
                sentences.append(f"Finally, {other_str} also increase the risk.")
            elif clinical_factors or situational_factors:
                # Only one previous section has content
                sentences.append(f"Also, {other_str} increase the risk.")
            else:
                # Neither previous section has content
                sentences.append(f"Risk is enhanced by {other_str}.")

        return " ".join(sentences)

    def _update_enhancing_narrative(self):
        """Update Risk-Enhancing Factors card from selections."""
        self._update_preview("risk_enhancing")

    def _build_popup_protective_factors(self):
        """Build Protective Factors popup distinguishing strong vs weak protectors."""
        container, layout = self._create_popup_container("protective")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("7. Protective Factors")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("Only include factors that operate in real life, not theoretical ones. Distinguish strong from weak/conditional protectors.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Strong protectors container
        strong_container = QFrame()
        strong_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(34, 197, 94, 0.4);
                border-radius: 10px;
            }
        """)
        strong_layout = QVBoxLayout(strong_container)
        strong_layout.setContentsMargins(12, 10, 12, 10)
        strong_layout.setSpacing(8)

        strong_lbl = QLabel("Strong Protectors")
        strong_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #166534; background: transparent; border: none;")
        strong_layout.addWidget(strong_lbl)

        STRONG_PROTECTORS = [
            ("prot_treatment_adherence", "Sustained treatment adherence"),
            ("prot_structured_supervision", "Structured supervision in place"),
            ("prot_supportive_relationships", "Supportive, prosocial relationships"),
            ("prot_insight_linked", "Insight linked to behaviour change"),
            ("prot_help_seeking", "Early help-seeking behaviour"),
            ("prot_restricted_access", "Restricted access to triggers/victims"),
            ("prot_medication_response", "Good response to medication"),
        ]

        self.popup_strong_protector_checkboxes = []
        for item_key, item_label in STRONG_PROTECTORS:
            cb = QCheckBox(item_label)
            cb.setProperty("prot_key", item_key)
            cb.setProperty("prot_strength", "strong")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_protective_narrative)
            strong_layout.addWidget(cb)
            self.popup_strong_protector_checkboxes.append(cb)

        layout.addWidget(strong_container)

        # Weak/conditional protectors container
        weak_container = QFrame()
        weak_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        weak_layout = QVBoxLayout(weak_container)
        weak_layout.setContentsMargins(12, 10, 12, 10)
        weak_layout.setSpacing(8)

        weak_lbl = QLabel("Weak / Conditional Protectors (flag as such)")
        weak_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        weak_layout.addWidget(weak_lbl)

        WEAK_PROTECTORS = [
            ("prot_verbal_motivation", "Verbal motivation only (untested)"),
            ("prot_untested_coping", "Untested coping skills"),
            ("prot_conditional_support", "Supports that may disengage under stress"),
            ("prot_external_motivation", "Externally motivated compliance only"),
            ("prot_situational_stability", "Stability dependent on current environment"),
        ]

        self.popup_weak_protector_checkboxes = []
        for item_key, item_label in WEAK_PROTECTORS:
            cb = QCheckBox(item_label)
            cb.setProperty("prot_key", item_key)
            cb.setProperty("prot_strength", "weak")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_protective_narrative)
            weak_layout.addWidget(cb)
            self.popup_weak_protector_checkboxes.append(cb)

        layout.addWidget(weak_container)

        layout.addStretch()

        def generate():
            return self._generate_protective_narrative()

        self._connect_preview_updates("protective", [])
        self._add_send_button(layout, "protective", generate)

    def _generate_protective_narrative(self):
        """Generate Protective Factors narrative from selections."""
        strong_phrases = {
            "prot_treatment_adherence": "sustained treatment adherence",
            "prot_structured_supervision": "structured supervision",
            "prot_supportive_relationships": "supportive, prosocial relationships",
            "prot_insight_linked": "insight linked to behaviour change",
            "prot_help_seeking": "early help-seeking behaviour",
            "prot_restricted_access": "restricted access to triggers and victims",
            "prot_medication_response": "demonstrated response to medication",
        }

        weak_phrases = {
            "prot_verbal_motivation": "verbal motivation (untested)",
            "prot_untested_coping": "untested coping skills",
            "prot_conditional_support": "supports that may disengage under stress",
            "prot_external_motivation": "externally motivated compliance only",
            "prot_situational_stability": "stability dependent on current environment",
        }

        strong = []
        weak = []

        for cb in getattr(self, 'popup_strong_protector_checkboxes', []):
            if cb.isChecked():
                key = cb.property("prot_key")
                if key in strong_phrases:
                    strong.append(strong_phrases[key])

        for cb in getattr(self, 'popup_weak_protector_checkboxes', []):
            if cb.isChecked():
                key = cb.property("prot_key")
                if key in weak_phrases:
                    weak.append(weak_phrases[key])

        parts = []
        if strong:
            strong_str = ", ".join(strong[:-1]) + " and " + strong[-1] if len(strong) > 1 else strong[0]
            parts.append(f"Protective factors include {strong_str}")

        if weak:
            weak_str = ", ".join(weak[:-1]) + " and " + weak[-1] if len(weak) > 1 else weak[0]
            if parts:
                parts.append(f"However, some protectors are weak or conditional, including {weak_str}")
            else:
                parts.append(f"Identified protectors are weak or conditional, including {weak_str}")

        return ". ".join(parts) + "." if parts else ""

    def _update_protective_narrative(self):
        """Update Protective Factors card from selections."""
        self._update_preview("protective")

    def _build_popup_monitoring(self):
        """Build Risk Monitoring Indicators popup with early warning signs."""
        container, layout = self._create_popup_container("monitoring")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("8. Risk Monitoring Indicators")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("Early warning signs that risk is increasing. These are gold for ongoing monitoring.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Behavioural indicators container
        behav_container = QFrame()
        behav_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        behav_layout = QVBoxLayout(behav_container)
        behav_layout.setContentsMargins(12, 10, 12, 10)
        behav_layout.setSpacing(8)

        behav_lbl = QLabel("Behavioural Indicators")
        behav_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        behav_layout.addWidget(behav_lbl)

        BEHAVIOURAL = [
            ("mon_missed_appts", "Missed appointments"),
            ("mon_med_refusal", "Medication refusal"),
            ("mon_withdrawal", "Withdrawal from supports"),
            ("mon_substance_use", "Increased substance use"),
            ("mon_non_compliance", "Non-compliance with conditions"),
            ("mon_rule_breaking", "Rule-breaking behaviour"),
        ]

        self.popup_behav_monitor_checkboxes = []
        for item_key, item_label in BEHAVIOURAL:
            cb = QCheckBox(item_label)
            cb.setProperty("mon_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_monitoring_narrative)
            behav_layout.addWidget(cb)
            self.popup_behav_monitor_checkboxes.append(cb)

        layout.addWidget(behav_container)

        # Mental state indicators container
        mental_container = QFrame()
        mental_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        mental_layout = QVBoxLayout(mental_container)
        mental_layout.setContentsMargins(12, 10, 12, 10)
        mental_layout.setSpacing(8)

        mental_lbl = QLabel("Mental State Indicators")
        mental_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        mental_layout.addWidget(mental_lbl)

        MENTAL_STATE = [
            ("mon_sleep_disturb", "Sleep disturbance"),
            ("mon_paranoia", "Rising paranoia or suspiciousness"),
            ("mon_irritability", "Increasing irritability"),
            ("mon_hostile_language", "Escalation in hostile language"),
            ("mon_fixation", "Fixation on grievances"),
            ("mon_agitation", "Increased agitation or restlessness"),
        ]

        self.popup_mental_monitor_checkboxes = []
        for item_key, item_label in MENTAL_STATE:
            cb = QCheckBox(item_label)
            cb.setProperty("mon_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_monitoring_narrative)
            mental_layout.addWidget(cb)
            self.popup_mental_monitor_checkboxes.append(cb)

        layout.addWidget(mental_container)

        layout.addStretch()

        def generate():
            return self._generate_monitoring_narrative()

        self._connect_preview_updates("monitoring", [])
        self._add_send_button(layout, "monitoring", generate)

    def _generate_monitoring_narrative(self):
        """Generate Monitoring Indicators narrative from selections."""
        monitor_phrases = {
            "mon_missed_appts": "missed appointments",
            "mon_med_refusal": "medication refusal",
            "mon_withdrawal": "withdrawal from supports",
            "mon_substance_use": "increased substance use",
            "mon_non_compliance": "non-compliance with conditions",
            "mon_rule_breaking": "rule-breaking behaviour",
            "mon_sleep_disturb": "sleep disturbance",
            "mon_paranoia": "rising paranoia or suspiciousness",
            "mon_irritability": "increasing irritability",
            "mon_hostile_language": "escalation in hostile language",
            "mon_fixation": "fixation on grievances",
            "mon_agitation": "increased agitation or restlessness",
        }

        indicators = []
        for cb_list in [
            getattr(self, 'popup_behav_monitor_checkboxes', []),
            getattr(self, 'popup_mental_monitor_checkboxes', []),
        ]:
            for cb in cb_list:
                if cb.isChecked():
                    key = cb.property("mon_key")
                    if key in monitor_phrases:
                        indicators.append(monitor_phrases[key])

        if indicators:
            ind_str = ", ".join(indicators[:-1]) + " and " + indicators[-1] if len(indicators) > 1 else indicators[0]
            return f"Early warning signs to monitor include {ind_str}."
        return ""

    def _update_monitoring_narrative(self):
        """Update Monitoring Indicators card from selections."""
        self._update_preview("monitoring")

    def _build_popup_management(self):
        """Build Risk Management Strategies popup with prevention, containment, response."""
        container, layout = self._create_popup_container("treatment")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("9. Risk Management Strategies")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("What actually helps? Split into prevention, containment, and response strategies.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Prevention container
        prevent_container = QFrame()
        prevent_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(34, 197, 94, 0.4);
                border-radius: 10px;
            }
        """)
        prevent_layout = QVBoxLayout(prevent_container)
        prevent_layout.setContentsMargins(12, 10, 12, 10)
        prevent_layout.setSpacing(8)

        prevent_lbl = QLabel("Preventative Strategies")
        prevent_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #166534; background: transparent; border: none;")
        prevent_layout.addWidget(prevent_lbl)

        PREVENTION = [
            ("mgmt_med_adherence", "Medication adherence"),
            ("mgmt_regular_review", "Regular clinical review"),
            ("mgmt_structured_routine", "Structured daily routine"),
            ("mgmt_stress_management", "Stress management interventions"),
            ("mgmt_substance_controls", "Substance use controls / monitoring"),
            ("mgmt_therapy", "Psychological therapy"),
        ]

        self.popup_prevent_checkboxes = []
        for item_key, item_label in PREVENTION:
            cb = QCheckBox(item_label)
            cb.setProperty("mgmt_key", item_key)
            cb.setProperty("mgmt_type", "prevent")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_management_narrative)
            prevent_layout.addWidget(cb)
            self.popup_prevent_checkboxes.append(cb)

        layout.addWidget(prevent_container)

        # Containment container
        contain_container = QFrame()
        contain_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        contain_layout = QVBoxLayout(contain_container)
        contain_layout.setContentsMargins(12, 10, 12, 10)
        contain_layout.setSpacing(8)

        contain_lbl = QLabel("Containment Strategies")
        contain_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        contain_layout.addWidget(contain_lbl)

        CONTAINMENT = [
            ("mgmt_supervision", "Ongoing supervision"),
            ("mgmt_conditions", "Conditions / boundaries"),
            ("mgmt_reduced_access", "Reduced access to triggers"),
            ("mgmt_supported_accom", "Supported accommodation"),
            ("mgmt_curfew", "Curfew or time restrictions"),
            ("mgmt_geographic", "Geographic restrictions"),
        ]

        self.popup_contain_checkboxes = []
        for item_key, item_label in CONTAINMENT:
            cb = QCheckBox(item_label)
            cb.setProperty("mgmt_key", item_key)
            cb.setProperty("mgmt_type", "contain")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_management_narrative)
            contain_layout.addWidget(cb)
            self.popup_contain_checkboxes.append(cb)

        layout.addWidget(contain_container)

        # Response container
        response_container = QFrame()
        response_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        response_layout = QVBoxLayout(response_container)
        response_layout.setContentsMargins(12, 10, 12, 10)
        response_layout.setSpacing(8)

        response_lbl = QLabel("Response Strategies")
        response_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        response_layout.addWidget(response_lbl)

        RESPONSE = [
            ("mgmt_escalation", "Clear escalation pathways"),
            ("mgmt_crisis_plan", "Crisis plan in place"),
            ("mgmt_recall_threshold", "Defined recall / admission thresholds"),
            ("mgmt_out_of_hours", "Out-of-hours response plan"),
            ("mgmt_police_protocol", "Police liaison protocol"),
        ]

        self.popup_response_checkboxes = []
        for item_key, item_label in RESPONSE:
            cb = QCheckBox(item_label)
            cb.setProperty("mgmt_key", item_key)
            cb.setProperty("mgmt_type", "response")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_management_narrative)
            response_layout.addWidget(cb)
            self.popup_response_checkboxes.append(cb)

        layout.addWidget(response_container)

        layout.addStretch()

        def generate():
            return self._generate_management_narrative()

        self._connect_preview_updates("treatment", [])
        self._add_send_button(layout, "treatment", generate)

    def _generate_management_narrative(self):
        """Generate Risk Management Strategies narrative from selections."""
        prevent_phrases = {
            "mgmt_med_adherence": "medication adherence",
            "mgmt_regular_review": "regular clinical review",
            "mgmt_structured_routine": "structured daily routine",
            "mgmt_stress_management": "stress management interventions",
            "mgmt_substance_controls": "substance use controls and monitoring",
            "mgmt_therapy": "psychological therapy",
        }

        contain_phrases = {
            "mgmt_supervision": "ongoing supervision",
            "mgmt_conditions": "appropriate conditions and boundaries",
            "mgmt_reduced_access": "reduced access to triggers",
            "mgmt_supported_accom": "supported accommodation",
            "mgmt_curfew": "curfew or time restrictions",
            "mgmt_geographic": "geographic restrictions",
        }

        response_phrases = {
            "mgmt_escalation": "clear escalation pathways",
            "mgmt_crisis_plan": "crisis plan",
            "mgmt_recall_threshold": "defined recall or admission thresholds",
            "mgmt_out_of_hours": "out-of-hours response plan",
            "mgmt_police_protocol": "police liaison protocol",
        }

        prevent = []
        contain = []
        response = []

        for cb in getattr(self, 'popup_prevent_checkboxes', []):
            if cb.isChecked():
                key = cb.property("mgmt_key")
                if key in prevent_phrases:
                    prevent.append(prevent_phrases[key])

        for cb in getattr(self, 'popup_contain_checkboxes', []):
            if cb.isChecked():
                key = cb.property("mgmt_key")
                if key in contain_phrases:
                    contain.append(contain_phrases[key])

        for cb in getattr(self, 'popup_response_checkboxes', []):
            if cb.isChecked():
                key = cb.property("mgmt_key")
                if key in response_phrases:
                    response.append(response_phrases[key])

        sentences = []

        # Preventative strategies sentence
        if prevent:
            prev_str = ", ".join(prevent[:-1]) + " and " + prevent[-1] if len(prevent) > 1 else prevent[0]
            sentences.append(f"Preventative strategies include {prev_str}.")

        # Containment sentence
        if contain:
            cont_str = ", ".join(contain[:-1]) + " and " + contain[-1] if len(contain) > 1 else contain[0]
            sentences.append(f"For containment we would recommend {cont_str}.")

        # Response sentence
        if response:
            resp_str = ", ".join(response[:-1]) + " and " + response[-1] if len(response) > 1 else response[0]
            sentences.append(f"A well formulated response to increased risk would involve {resp_str}.")

        return " ".join(sentences)

    def _update_management_narrative(self):
        """Update Risk Management Strategies card from selections."""
        self._update_preview("treatment")

    def _build_popup_supervision_recs(self):
        """Build Supervision Recommendations popup with specific recommendations."""
        container, layout = self._create_popup_container("supervision")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"
        radio_style = "QRadioButton { font-size: 15px; padding: 4px; }"

        lbl = QLabel("10. Supervision Recommendations")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("Be specific about level, frequency, who monitors what, and escalation triggers.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Supervision level container
        level_container = QFrame()
        level_container.setStyleSheet("""
            QFrame {
                background: rgba(219, 234, 254, 0.5);
                border: 2px solid rgba(59, 130, 246, 0.4);
                border-radius: 10px;
            }
        """)
        level_layout = QVBoxLayout(level_container)
        level_layout.setContentsMargins(12, 10, 12, 10)
        level_layout.setSpacing(8)

        level_lbl = QLabel("Supervision Level")
        level_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #1d4ed8; background: transparent; border: none;")
        level_layout.addWidget(level_lbl)

        self.popup_supervision_level_group = QButtonGroup(self)

        sup_informal = QRadioButton("Informal — voluntary engagement")
        sup_informal.setStyleSheet(radio_style)
        self.popup_supervision_level_group.addButton(sup_informal, 0)
        level_layout.addWidget(sup_informal)

        sup_supported = QRadioButton("Supported — structured community support")
        sup_supported.setStyleSheet(radio_style)
        self.popup_supervision_level_group.addButton(sup_supported, 1)
        level_layout.addWidget(sup_supported)

        sup_conditional = QRadioButton("Conditional — with enforceable conditions")
        sup_conditional.setStyleSheet(radio_style)
        self.popup_supervision_level_group.addButton(sup_conditional, 2)
        level_layout.addWidget(sup_conditional)

        sup_restricted = QRadioButton("Restricted — secure or hospital setting")
        sup_restricted.setStyleSheet(radio_style)
        self.popup_supervision_level_group.addButton(sup_restricted, 3)
        level_layout.addWidget(sup_restricted)

        self.popup_supervision_level_group.buttonClicked.connect(self._update_supervision_narrative)
        layout.addWidget(level_container)

        # Contact frequency container
        contact_container = QFrame()
        contact_container.setStyleSheet("""
            QFrame {
                background: rgba(220, 252, 231, 0.5);
                border: 2px solid rgba(34, 197, 94, 0.4);
                border-radius: 10px;
            }
        """)
        contact_layout = QVBoxLayout(contact_container)
        contact_layout.setContentsMargins(12, 10, 12, 10)
        contact_layout.setSpacing(8)

        contact_lbl = QLabel("Contact Requirements")
        contact_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #166534; background: transparent; border: none;")
        contact_layout.addWidget(contact_lbl)

        CONTACT_REQS = [
            ("sup_face_to_face", "Regular face-to-face contact required"),
            ("sup_med_monitoring", "Medication adherence monitoring"),
            ("sup_urine_screening", "Urine screening for substances"),
            ("sup_curfew_checks", "Curfew checks"),
            ("sup_unannounced", "Unannounced visits"),
            ("sup_phone_checkins", "Regular phone check-ins"),
        ]

        self.popup_contact_checkboxes = []
        for item_key, item_label in CONTACT_REQS:
            cb = QCheckBox(item_label)
            cb.setProperty("sup_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_supervision_narrative)
            contact_layout.addWidget(cb)
            self.popup_contact_checkboxes.append(cb)

        layout.addWidget(contact_container)

        # Escalation triggers container
        escalate_container = QFrame()
        escalate_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        escalate_layout = QVBoxLayout(escalate_container)
        escalate_layout.setContentsMargins(12, 10, 12, 10)
        escalate_layout.setSpacing(8)

        escalate_lbl = QLabel("Escalation Triggers")
        escalate_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        escalate_layout.addWidget(escalate_lbl)

        ESCALATION = [
            ("esc_engagement_deteriorates", "Engagement deteriorates"),
            ("esc_non_compliance", "Non-compliance with conditions"),
            ("esc_warning_signs", "Early warning signs emerge"),
            ("esc_substance_relapse", "Substance use relapse"),
            ("esc_mental_state", "Mental state deterioration"),
            ("esc_threats", "Threats or aggressive behaviour"),
        ]

        self.popup_escalation_checkboxes = []
        for item_key, item_label in ESCALATION:
            cb = QCheckBox(item_label)
            cb.setProperty("esc_key", item_key)
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_supervision_narrative)
            escalate_layout.addWidget(cb)
            self.popup_escalation_checkboxes.append(cb)

        layout.addWidget(escalate_container)

        layout.addStretch()

        def generate():
            return self._generate_supervision_narrative()

        self._connect_preview_updates("supervision", [])
        self._add_send_button(layout, "supervision", generate)

    def _generate_supervision_narrative(self):
        """Generate Supervision Recommendations narrative from selections."""
        sentences = []

        # Supervision level - first sentence
        btn = self.popup_supervision_level_group.checkedButton()
        level_text = ""
        if btn:
            btn_id = self.popup_supervision_level_group.id(btn)
            if btn_id == 0:
                level_text = "Ongoing supervision at an informal level with voluntary engagement is recommended."
            elif btn_id == 1:
                level_text = "Ongoing supervision with structured community support is recommended."
            elif btn_id == 2:
                level_text = "Ongoing supervision with enforceable conditions is recommended."
            elif btn_id == 3:
                level_text = "Continued management in a secure or hospital setting is recommended."
            sentences.append(level_text)

        contact_phrases = {
            "sup_face_to_face": "regular face-to-face contact",
            "sup_med_monitoring": "monitoring of medication adherence",
            "sup_urine_screening": "urine screening for substances",
            "sup_curfew_checks": "curfew checks",
            "sup_unannounced": "unannounced visits",
            "sup_phone_checkins": "regular phone check-ins",
        }

        contacts = []
        for cb in getattr(self, 'popup_contact_checkboxes', []):
            if cb.isChecked():
                key = cb.property("sup_key")
                if key in contact_phrases:
                    contacts.append(contact_phrases[key])

        esc_phrases = {
            "esc_engagement_deteriorates": "engagement deteriorates",
            "esc_non_compliance": "non-compliance with conditions",
            "esc_warning_signs": "early warning signs emerge",
            "esc_substance_relapse": "substance use relapse",
            "esc_mental_state": "mental state deterioration",
            "esc_threats": "threats or aggressive behaviour",
        }

        escalations = []
        for cb in getattr(self, 'popup_escalation_checkboxes', []):
            if cb.isChecked():
                key = cb.property("esc_key")
                if key in esc_phrases:
                    escalations.append(esc_phrases[key])

        # Build second sentence combining contacts and escalation triggers
        if contacts or escalations:
            second_sentence_parts = []

            if contacts:
                cont_str = ", ".join(contacts[:-1]) + " and " + contacts[-1] if len(contacts) > 1 else contacts[0]
                if level_text:
                    second_sentence_parts.append(f"This should include {cont_str}")
                else:
                    second_sentence_parts.append(f"Supervision should include {cont_str}")

            if escalations:
                esc_str = ", ".join(escalations[:-1]) + " or " + escalations[-1] if len(escalations) > 1 else escalations[0]
                if second_sentence_parts:
                    second_sentence_parts.append(f"with prompt review if {esc_str}")
                else:
                    second_sentence_parts.append(f"Prompt review is required if {esc_str}")

            sentences.append(" ".join(second_sentence_parts) + ".")

        narrative = " ".join(sentences)
        return narrative.replace("..", ".")

    def _update_supervision_narrative(self):
        """Update Supervision Recommendations card from selections."""
        self._update_preview("supervision")

    def _build_popup_victim_safety(self):
        """Build Victim Safety & Safeguarding Planning popup."""
        container, layout = self._create_popup_container("victim_safety")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        lbl = QLabel("11. Victim Safety & Safeguarding Planning")
        lbl.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        layout.addWidget(lbl)

        hint = QLabel("This must be explicit, even if no named victim exists. Select applicable strategies.")
        hint.setStyleSheet("font-size: 15px; color: #6b7280; font-style: italic;")
        hint.setWordWrap(True)
        layout.addWidget(hint)

        # Named victim container
        named_container = QFrame()
        named_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 226, 226, 0.5);
                border: 2px solid rgba(220, 38, 38, 0.4);
                border-radius: 10px;
            }
        """)
        named_layout = QVBoxLayout(named_container)
        named_layout.setContentsMargins(12, 10, 12, 10)
        named_layout.setSpacing(8)

        named_lbl = QLabel("When There Is an Identified Victim")
        named_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #991b1b; background: transparent; border: none;")
        named_layout.addWidget(named_lbl)

        NAMED_VICTIM = [
            ("vic_separation", "Physical separation maintained"),
            ("vic_no_contact", "No-contact conditions in place"),
            ("vic_third_party", "Third-party monitoring"),
            ("vic_info_sharing", "Information sharing between agencies"),
            ("vic_victim_informed", "Victim informed of risk and release"),
            ("vic_exclusion_zone", "Exclusion zone in place"),
        ]

        self.popup_named_victim_checkboxes = []
        for item_key, item_label in NAMED_VICTIM:
            cb = QCheckBox(item_label)
            cb.setProperty("vic_key", item_key)
            cb.setProperty("vic_type", "named")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_victim_safety_narrative)
            named_layout.addWidget(cb)
            self.popup_named_victim_checkboxes.append(cb)

        layout.addWidget(named_container)

        # Non-specific victim container
        general_container = QFrame()
        general_container.setStyleSheet("""
            QFrame {
                background: rgba(254, 243, 199, 0.5);
                border: 2px solid rgba(217, 119, 6, 0.4);
                border-radius: 10px;
            }
        """)
        general_layout = QVBoxLayout(general_container)
        general_layout.setContentsMargins(12, 10, 12, 10)
        general_layout.setSpacing(8)

        general_lbl = QLabel("When Victims Are Non-Specific")
        general_lbl.setStyleSheet("font-size: 16px; font-weight: 600; color: #92400e; background: transparent; border: none;")
        general_layout.addWidget(general_lbl)

        GENERAL_SAFETY = [
            ("vic_env_controls", "Environmental controls"),
            ("vic_staff_safety", "Staff safety planning"),
            ("vic_conflict_avoid", "Conflict avoidance strategies"),
            ("vic_de_escalation", "De-escalation protocols"),
            ("vic_restricted_access", "Restricted access to vulnerable groups"),
            ("vic_public_protection", "Public protection measures"),
        ]

        self.popup_general_safety_checkboxes = []
        for item_key, item_label in GENERAL_SAFETY:
            cb = QCheckBox(item_label)
            cb.setProperty("vic_key", item_key)
            cb.setProperty("vic_type", "general")
            cb.setStyleSheet("QCheckBox { background: transparent; border: none; font-size: 14px; color: #374151; padding: 2px; }")
            cb.stateChanged.connect(self._update_victim_safety_narrative)
            general_layout.addWidget(cb)
            self.popup_general_safety_checkboxes.append(cb)

        layout.addWidget(general_container)

        layout.addStretch()

        def generate():
            return self._generate_victim_safety_narrative()

        self._connect_preview_updates("victim_safety", [])
        self._add_send_button(layout, "victim_safety", generate)

    def _generate_victim_safety_narrative(self):
        """Generate Victim Safety narrative from selections."""
        named_phrases = {
            "vic_separation": "physical separation from previous victims",
            "vic_no_contact": "no-contact conditions",
            "vic_third_party": "third-party monitoring",
            "vic_info_sharing": "information sharing between agencies",
            "vic_victim_informed": "victim notification of risk and release",
            "vic_exclusion_zone": "exclusion zone around victim locations",
        }

        general_phrases = {
            "vic_env_controls": "environmental controls",
            "vic_staff_safety": "staff safety planning",
            "vic_conflict_avoid": "conflict avoidance strategies",
            "vic_de_escalation": "de-escalation protocols",
            "vic_restricted_access": "restricted access to vulnerable groups",
            "vic_public_protection": "public protection measures",
        }

        named = []
        general = []

        for cb in getattr(self, 'popup_named_victim_checkboxes', []):
            if cb.isChecked():
                key = cb.property("vic_key")
                if key in named_phrases:
                    named.append(named_phrases[key])

        for cb in getattr(self, 'popup_general_safety_checkboxes', []):
            if cb.isChecked():
                key = cb.property("vic_key")
                if key in general_phrases:
                    general.append(general_phrases[key])

        parts = []
        if named:
            named_str = ", ".join(named[:-1]) + " and " + named[-1] if len(named) > 1 else named[0]
            parts.append(f"Victim safety planning should include {named_str}")

        if general:
            gen_str = ", ".join(general[:-1]) + " and " + general[-1] if len(general) > 1 else general[0]
            if parts:
                parts.append(f"General safeguarding includes {gen_str}")
            else:
                parts.append(f"Safeguarding planning should include {gen_str}")

        return ". ".join(parts) + "." if parts else ""

    def _update_victim_safety_narrative(self):
        """Update Victim Safety card from selections."""
        self._update_preview("victim_safety")

    def _auto_populate_risk_formulation(self):
        """Auto-populate risk formulation checkboxes based on H1-R5 item selections."""
        # This method checks the state of H1-R5 checkboxes and auto-selects
        # relevant checkboxes in the risk formulation section

        # === Risk-Enhancing Factors auto-population from C1-C5 ===

        # C1 - Insight: if poor insight selected, auto-check enhancer
        c1_disorder_checkboxes = getattr(self, 'popup_c1_disorder_checkboxes', [])
        for cb in c1_disorder_checkboxes:
            if cb.isChecked() and cb.property("c1_key") in ["dis_denies_illness", "dis_rejects_diagnosis", "dis_poor_insight"]:
                # Auto-check poor insight enhancer
                for enh_cb in getattr(self, 'popup_clinical_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_poor_insight":
                        enh_cb.setChecked(True)
                break

        # C2 - Violent Ideation: if any violent ideation selected
        c2_checkboxes = getattr(self, 'popup_c2_explicit_checkboxes', []) + getattr(self, 'popup_c2_threats_checkboxes', [])
        for cb in c2_checkboxes:
            if cb.isChecked():
                for enh_cb in getattr(self, 'popup_clinical_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_violent_ideation":
                        enh_cb.setChecked(True)
                break

        # C3 - Active Symptoms: if any psychotic/manic symptoms selected
        c3_psychotic = getattr(self, 'popup_c3_psychotic_checkboxes', [])
        c3_mania = getattr(self, 'popup_c3_mania_checkboxes', [])
        for cb in c3_psychotic + c3_mania:
            if cb.isChecked():
                for enh_cb in getattr(self, 'popup_clinical_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_active_symptoms":
                        enh_cb.setChecked(True)
                break

        # C4 - Instability: if instability indicators selected
        c4_checkboxes = getattr(self, 'popup_c4_affective_checkboxes', []) + getattr(self, 'popup_c4_behav_checkboxes', [])
        for cb in c4_checkboxes:
            if cb.isChecked():
                for enh_cb in getattr(self, 'popup_clinical_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_instability":
                        enh_cb.setChecked(True)
                break

        # C5 - Treatment Response: if poor treatment response selected
        c5_checkboxes = getattr(self, 'popup_c5_adherence_checkboxes', []) + getattr(self, 'popup_c5_engagement_checkboxes', [])
        for cb in c5_checkboxes:
            key = cb.property("c5_key") if cb.property("c5_key") else ""
            if cb.isChecked() and "non" in key.lower() or "refus" in key.lower() or "poor" in key.lower():
                for enh_cb in getattr(self, 'popup_clinical_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_poor_treatment":
                        enh_cb.setChecked(True)
                break

        # === Situational enhancers from R1-R5 ===

        # R1 - Professional Services: if poor plan indicators
        r1_checkboxes = getattr(self, 'popup_r1_plan_checkboxes', [])
        for cb in r1_checkboxes:
            key = cb.property("r1_key") if cb.property("r1_key") else ""
            if cb.isChecked() and ("no_clear" in key or "inadequate" in key or "incomplete" in key):
                for enh_cb in getattr(self, 'popup_situational_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_poor_plan":
                        enh_cb.setChecked(True)
                break

        # R2 - Living Situation: if unstable accommodation
        r2_checkboxes = getattr(self, 'popup_r2_accom_checkboxes', [])
        for cb in r2_checkboxes:
            key = cb.property("r2_key") if cb.property("r2_key") else ""
            if cb.isChecked() and ("unstable" in key or "homeless" in key or "temporary" in key):
                for enh_cb in getattr(self, 'popup_situational_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_unstable_living":
                        enh_cb.setChecked(True)
                break

        # R3 - Personal Support: if poor support
        r3_checkboxes = getattr(self, 'popup_r3_support_checkboxes', [])
        for cb in r3_checkboxes:
            key = cb.property("r3_key") if cb.property("r3_key") else ""
            if cb.isChecked() and ("isolated" in key or "no_support" in key or "limited" in key):
                for enh_cb in getattr(self, 'popup_situational_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_poor_support":
                        enh_cb.setChecked(True)
                break

        # R4 - Compliance: if non-compliance
        r4_checkboxes = getattr(self, 'popup_r4_adherence_checkboxes', []) + getattr(self, 'popup_r4_attendance_checkboxes', [])
        for cb in r4_checkboxes:
            key = cb.property("r4_key") if cb.property("r4_key") else ""
            if cb.isChecked() and ("non" in key or "refus" in key or "miss" in key):
                for enh_cb in getattr(self, 'popup_situational_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_non_compliance":
                        enh_cb.setChecked(True)
                break

        # R5 - Stress/Coping: if poor coping
        r5_checkboxes = getattr(self, 'popup_r5_coping_checkboxes', [])
        for cb in r5_checkboxes:
            key = cb.property("r5_key") if cb.property("r5_key") else ""
            if cb.isChecked() and ("poor" in key or "limited" in key or "maladaptive" in key):
                for enh_cb in getattr(self, 'popup_situational_enhancer_checkboxes', []):
                    if enh_cb.property("enh_key") == "enh_poor_coping":
                        enh_cb.setChecked(True)
                break

        # Update narratives after auto-population
        self._update_enhancing_narrative()

    # Legacy compatibility wrappers for old popup builders
    def _build_popup_scenario(self, key: str, title: str, prompt: str):
        """Legacy wrapper - redirects to new enhanced popup builders."""
        # This is now handled by the specific popup builders above
        pass

    def _build_popup_list_section(self, key: str, title: str, prompt: str):
        """Legacy wrapper - redirects to new enhanced popup builders."""
        # This is now handled by the specific popup builders above
        pass

    def _build_popup_signature(self):
        """Build signature popup."""
        container, layout = self._create_popup_container("signature")

        label_style = "font-size: 17px; color: #374151;"
        input_style = "padding: 8px; border: 1px solid #d1d5db; border-radius: 4px; font-size: 17px;"

        # Author name - auto-fill from MyDetails
        name_lbl = QLabel("Author Name:")
        name_lbl.setStyleSheet(label_style)
        layout.addWidget(name_lbl)
        self.popup_sig_name = QLineEdit()
        self.popup_sig_name.setStyleSheet(input_style)
        # Auto-fill from MyDetails
        if self._my_details.get("full_name"):
            self.popup_sig_name.setText(self._my_details["full_name"])
        layout.addWidget(self.popup_sig_name)

        # Role/Title - auto-fill from MyDetails
        role_lbl = QLabel("Role/Title:")
        role_lbl.setStyleSheet(label_style)
        layout.addWidget(role_lbl)
        self.popup_sig_role = QLineEdit()
        self.popup_sig_role.setPlaceholderText("e.g., Principal Forensic Psychologist")
        self.popup_sig_role.setStyleSheet(input_style)
        # Auto-fill from MyDetails
        if self._my_details.get("role_title"):
            self.popup_sig_role.setText(self._my_details["role_title"])
        layout.addWidget(self.popup_sig_role)

        # Date
        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet(label_style)
        layout.addWidget(date_lbl)
        self.popup_sig_date = QDateEdit()
        self.popup_sig_date.setCalendarPopup(True)
        self.popup_sig_date.setDisplayFormat("dd.MM.yy")
        self.popup_sig_date.setDate(QDate.currentDate())
        self.popup_sig_date.setStyleSheet(input_style)
        layout.addWidget(self.popup_sig_date)

        layout.addStretch()

        def generate():
            parts = []
            name = self.popup_sig_name.text().strip()
            if name:
                parts.append(name)
            role = self.popup_sig_role.text().strip()
            if role:
                parts.append(role)
            date = self.popup_sig_date.date().toString("dd.MM.yy")
            parts.append(date)
            return "\n".join(parts)

        self._connect_preview_updates("signature", [
            self.popup_sig_name, self.popup_sig_role, self.popup_sig_date
        ])
        self._add_send_button(layout, "signature", generate)

    # ================================================================
    # ACTIONS
    # ================================================================

    def _clear_form(self):
        """Clear all form fields."""
        reply = QMessageBox.question(
            self, "Clear Form",
            "Are you sure you want to clear all fields?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            for key, card in self.cards.items():
                card.editor.clear()
            self._last_generated_text.clear()
            if hasattr(self, '_imported_report_data'):
                self._imported_report_data = {}
            if hasattr(self, '_imported_report_sections'):
                self._imported_report_sections = {}

            # Destroy all popups and remove from stack
            for key, popup in list(self.popups.items()):
                if hasattr(self, 'popup_stack'):
                    self.popup_stack.removeWidget(popup)
                popup.deleteLater()
            self.popups.clear()

            # Reset non-popup widgets (except signature name/role from my details)
            preserve = set()
            if hasattr(self, 'popup_sig_name'):
                preserve.add(self.popup_sig_name)
            if hasattr(self, 'popup_sig_role'):
                preserve.add(self.popup_sig_role)
            for widget in self.findChildren(QLineEdit):
                if widget not in preserve:
                    widget.clear()
            for widget in self.findChildren(QCheckBox):
                widget.setChecked(False)
            for widget in self.findChildren(QComboBox):
                widget.setCurrentIndex(0)
            for widget in self.findChildren(QDateEdit):
                widget.setDate(QDate.currentDate())
            if hasattr(self, 'popup_gender_male'):
                self.popup_gender_male.setChecked(False)
            if hasattr(self, 'popup_gender_female'):
                self.popup_gender_female.setChecked(False)

            # Restore signature card from my details
            if hasattr(self, 'popup_sig_name') and hasattr(self, 'popup_sig_role'):
                parts = []
                name = self.popup_sig_name.text().strip()
                if name:
                    parts.append(name)
                role = self.popup_sig_role.text().strip()
                if role:
                    parts.append(role)
                if hasattr(self, 'popup_sig_date'):
                    parts.append(self.popup_sig_date.date().toString("dd.MM.yy"))
                if parts and "signature" in self.cards:
                    self.cards["signature"].editor.setPlainText("\n".join(parts))

    # ================================================================
    # TOOLBAR FORMATTING METHODS
    # ================================================================

    def _get_active_editor(self):
        """Get the currently focused text editor."""
        # Use the tracked active editor (set via focus events)
        if self._active_editor and isinstance(self._active_editor, QTextEdit):
            return self._active_editor
        # Fallback: check if any QTextEdit currently has focus
        focused = QApplication.focusWidget()
        if isinstance(focused, QTextEdit):
            self._active_editor = focused
            return focused
        # Fallback: use the selected card's editor
        if self._selected_card_key and self._selected_card_key in self.cards:
            card = self.cards[self._selected_card_key]
            if hasattr(card, 'editor') and isinstance(card.editor, QTextEdit):
                return card.editor
        return None

    def _track_editor_focus(self, editor: QTextEdit):
        """Set up focus tracking for an editor."""
        # Use focusInEvent to track when this editor gets focus
        original_focus_in = editor.focusInEvent
        def new_focus_in(event):
            self._active_editor = editor
            original_focus_in(event)
        editor.focusInEvent = new_focus_in

    def _setup_editor_tracking(self):
        """Set up focus tracking for all card and popup editors."""
        # Track focus on all card editors
        for key, card in self.cards.items():
            if hasattr(card, 'editor') and isinstance(card.editor, QTextEdit):
                self._track_editor_focus(card.editor)

        # Track focus on all popup text editors
        for attr_name in dir(self):
            if attr_name.startswith('popup_') and attr_name.endswith('_text'):
                widget = getattr(self, attr_name, None)
                if isinstance(widget, QTextEdit):
                    self._track_editor_focus(widget)

        # Also track any QTextEdit children in popups
        for key, popup in self.popups.items():
            for child in popup.findChildren(QTextEdit):
                self._track_editor_focus(child)

    def _set_font_family(self, family: str):
        """Set font family on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setFontFamily(family)
            editor.mergeCurrentCharFormat(fmt)

    def _set_font_size(self, size: int):
        """Set font size on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setFontPointSize(size)
            editor.mergeCurrentCharFormat(fmt)

    def _toggle_bold(self):
        """Toggle bold on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setFontWeight(QFont.Weight.Normal if fmt.fontWeight() == QFont.Weight.Bold else QFont.Weight.Bold)
            editor.mergeCurrentCharFormat(fmt)

    def _toggle_italic(self):
        """Toggle italic on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setFontItalic(not fmt.fontItalic())
            editor.mergeCurrentCharFormat(fmt)

    def _toggle_underline(self):
        """Toggle underline on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setFontUnderline(not fmt.fontUnderline())
            editor.mergeCurrentCharFormat(fmt)

    def _set_text_color(self, color: QColor):
        """Set text color on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setForeground(color)
            editor.mergeCurrentCharFormat(fmt)

    def _set_highlight_color(self, color: QColor):
        """Set highlight/background color on selected text."""
        editor = self._get_active_editor()
        if editor:
            fmt = editor.currentCharFormat()
            fmt.setBackground(color)
            editor.mergeCurrentCharFormat(fmt)

    def _set_align_left(self):
        """Set left alignment."""
        editor = self._get_active_editor()
        if editor:
            editor.setAlignment(Qt.AlignmentFlag.AlignLeft)

    def _set_align_center(self):
        """Set center alignment."""
        editor = self._get_active_editor()
        if editor:
            editor.setAlignment(Qt.AlignmentFlag.AlignCenter)

    def _set_align_right(self):
        """Set right alignment."""
        editor = self._get_active_editor()
        if editor:
            editor.setAlignment(Qt.AlignmentFlag.AlignRight)

    def _set_align_justify(self):
        """Set justify alignment."""
        editor = self._get_active_editor()
        if editor:
            editor.setAlignment(Qt.AlignmentFlag.AlignJustify)

    def _toggle_bullet_list(self):
        """Toggle bullet list."""
        editor = self._get_active_editor()
        if editor:
            cursor = editor.textCursor()
            list_fmt = cursor.currentList()
            if list_fmt:
                # Remove list
                block_fmt = cursor.blockFormat()
                block_fmt.setIndent(0)
                cursor.setBlockFormat(block_fmt)
            else:
                # Create bullet list
                list_format = QTextListFormat()
                list_format.setStyle(QTextListFormat.Style.ListDisc)
                cursor.createList(list_format)

    def _toggle_numbered_list(self):
        """Toggle numbered list."""
        editor = self._get_active_editor()
        if editor:
            cursor = editor.textCursor()
            list_fmt = cursor.currentList()
            if list_fmt:
                # Remove list
                block_fmt = cursor.blockFormat()
                block_fmt.setIndent(0)
                cursor.setBlockFormat(block_fmt)
            else:
                # Create numbered list
                list_format = QTextListFormat()
                list_format.setStyle(QTextListFormat.Style.ListDecimal)
                cursor.createList(list_format)

    def _indent(self):
        """Increase indentation."""
        editor = self._get_active_editor()
        if editor:
            cursor = editor.textCursor()
            block_fmt = cursor.blockFormat()
            block_fmt.setIndent(block_fmt.indent() + 1)
            cursor.setBlockFormat(block_fmt)

    def _outdent(self):
        """Decrease indentation."""
        editor = self._get_active_editor()
        if editor:
            cursor = editor.textCursor()
            block_fmt = cursor.blockFormat()
            if block_fmt.indent() > 0:
                block_fmt.setIndent(block_fmt.indent() - 1)
                cursor.setBlockFormat(block_fmt)

    def _undo(self):
        """Undo last action."""
        editor = self._get_active_editor()
        if editor:
            editor.undo()

    def _redo(self):
        """Redo last undone action."""
        editor = self._get_active_editor()
        if editor:
            editor.redo()

    def _insert_date(self):
        """Insert current date at cursor."""
        editor = self._get_active_editor()
        if editor:
            from datetime import datetime
            date_str = datetime.now().strftime("%d %B %Y")
            editor.insertPlainText(date_str)

    def _export_docx(self):
        """Export to Word document matching the official HCR-20 template format."""
        from hcr20_docx_exporter import export_hcr20_docx

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export HCR-20 Report", "", "Word Documents (*.docx)"
        )
        if not file_path:
            return

        # Gather all data from the form
        data = self._collect_form_data()

        try:
            if export_hcr20_docx(data, file_path):
                QMessageBox.information(self, "Export Complete", f"Report exported to:\n{file_path}")
            else:
                QMessageBox.critical(self, "Export Error", "Failed to export the document.")
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export: {e}")

    def _collect_form_data(self) -> dict:
        """Collect all form data into a dictionary for export."""
        data = {}

        # Patient details
        if hasattr(self, 'popup_patient_name'):
            data['patient_name'] = self.popup_patient_name.text().strip()
        if hasattr(self, 'popup_dob'):
            data['dob'] = self.popup_dob.date().toString("d MMMM yyyy")
            # Calculate age
            from PySide6.QtCore import QDate
            today = QDate.currentDate()
            age = today.year() - self.popup_dob.date().year()
            if today.month() < self.popup_dob.date().month() or \
               (today.month() == self.popup_dob.date().month() and today.day() < self.popup_dob.date().day()):
                age -= 1
            data['age'] = str(age)
        if hasattr(self, 'popup_nhs'):
            data['nhs_number'] = self.popup_nhs.text().strip()
        if hasattr(self, 'popup_address'):
            data['address'] = self.popup_address.toPlainText().strip()
        if hasattr(self, 'popup_admission_date'):
            data['admission_date'] = self.popup_admission_date.date().toString("d MMMM yyyy")
        if hasattr(self, 'popup_legal_status'):
            data['legal_status'] = self.popup_legal_status.text().strip()

        # Assessment details
        if hasattr(self, 'popup_author_original'):
            data['author_original'] = self.popup_author_original.text().strip()
        if hasattr(self, 'popup_author_update'):
            data['author_update'] = self.popup_author_update.text().strip()
        if hasattr(self, 'popup_supervisor'):
            data['supervisor'] = self.popup_supervisor.text().strip()
        if hasattr(self, 'popup_review_to'):
            data['review_to'] = self.popup_review_to.text().strip()
        if hasattr(self, 'popup_date_original'):
            data['date_original'] = self.popup_date_original.date().toString("MMMM yyyy")
        if hasattr(self, 'popup_date_update'):
            data['date_update'] = self.popup_date_update.date().toString("MMMM yyyy")
        if hasattr(self, 'popup_date_next'):
            data['date_next'] = self.popup_date_next.date().toString("MMMM yyyy")

        # Sources
        if hasattr(self, 'popup_sources_text'):
            data['sources'] = self.popup_sources_text.toPlainText().strip()

        # HCR-20 Items (H1-H10, C1-C5, R1-R5)
        all_items = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'h7', 'h8', 'h9', 'h10',
                     'c1', 'c2', 'c3', 'c4', 'c5',
                     'r1', 'r2', 'r3', 'r4', 'r5']

        for key in all_items:
            item_data = {}

            # Get content from card editor
            if key in self.cards:
                item_data['content'] = self.cards[key].editor.toPlainText().strip()

            # Get presence rating
            presence_group = getattr(self, f"popup_{key}_presence_group", None)
            if presence_group:
                checked = presence_group.checkedButton()
                if checked:
                    text = checked.text()
                    if "No" in text:
                        item_data['presence'] = "Absent"
                    elif "Partial" in text:
                        item_data['presence'] = "Partially Present"
                    elif "Yes" in text:
                        item_data['presence'] = "Present"
                    elif "Omit" in text:
                        item_data['presence'] = "Omitted"

            # Get relevance rating
            relevance_group = getattr(self, f"popup_{key}_relevance_group", None)
            if relevance_group:
                checked = relevance_group.checkedButton()
                if checked:
                    text = checked.text()
                    if "Low" in text:
                        item_data['relevance'] = "Low relevance"
                    elif "Moderate" in text:
                        item_data['relevance'] = "Moderate relevance"
                    elif "High" in text:
                        item_data['relevance'] = "High relevance"

            data[key] = item_data

        # Formulation
        if hasattr(self, 'popup_formulation_text'):
            data['formulation'] = self.popup_formulation_text.toPlainText().strip()

        # Scenarios
        scenario_keys = ['scenario_nature', 'scenario_severity', 'scenario_imminence',
                        'scenario_frequency', 'scenario_likelihood']
        for key in scenario_keys:
            text_widget = getattr(self, f"popup_{key}_text", None)
            if text_widget:
                data[key] = text_widget.toPlainText().strip()

        # Management sections
        management_keys = ['risk_enhancing', 'protective', 'monitoring',
                          'treatment', 'supervision', 'victim_safety']
        for key in management_keys:
            text_widget = getattr(self, f"popup_{key}_text", None)
            if text_widget:
                data[key] = text_widget.toPlainText().strip()

        # Signature
        if hasattr(self, 'popup_sig_name'):
            data['signature_name'] = self.popup_sig_name.text().strip()
        if hasattr(self, 'popup_sig_role'):
            data['signature_role'] = self.popup_sig_role.text().strip()
        if hasattr(self, 'popup_sig_date'):
            data['signature_date'] = self.popup_sig_date.date().toString("dd.MM.yy")

        return data

    def _split_content_into_subsections(self, item_key: str, content: str) -> dict:
        """
        Split item content into subsections based on known header patterns.
        Returns a dict mapping subsection field names to their content.
        """
        import re

        # Define subsection header patterns for each item type
        # Maps item key -> list of (pattern_regex, field_suffix)
        subsection_patterns = {
            'h1': [
                (r'Child\s*(?:\(?\s*aged?\s*)?12\s*and\s*under\)?[:\s]*', 'child_(aged_12_and_under)'),
                (r'Adolescent\s*(?:\(?\s*aged?\s*)?(?:between\s*)?13\s*[-–]\s*17\)?[:\s]*', 'adolescent_(aged_13_17)'),
                (r'Adult\s*(?:\(?\s*aged?\s*)?18\s*(?:years\s*)?(?:and\s*over|over|\+)\)?[:\s]*', 'adult_(aged_18+)'),
            ],
            'h2': [
                (r'Child\s*(?:\(?\s*aged?\s*)?12\s*and\s*under\)?[:\s]*', 'child_(aged_12_and_under)'),
                (r'Adolescent\s*(?:\(?\s*aged?\s*)?(?:between\s*)?13\s*[-–]\s*17\)?[:\s]*', 'adolescent_(aged_13_17)'),
                (r'Adult\s*(?:\(?\s*aged?\s*)?18\s*(?:years\s*)?(?:and\s*over|over|\+)\)?[:\s]*', 'adult_(aged_18+)'),
            ],
            'h3': [
                (r'(?:^|\n)Intimate\s*[Rr]elationships?\s*\n', 'intimate_relationships'),
                (r'(?:^|\n)Non[-\s]*intimate\s*[Rr]elationships?\s*\n', 'non_intimate_relationships'),
            ],
            'h4': [
                (r'Education[:\s]*', 'education'),
                (r'Employment[:\s]*', 'employment'),
            ],
            'h5': [
                (r'Substance\s*Use\s*History[:\s]*', 'substance_use_history'),
                (r'Treatment\s*History[:\s]*', 'treatment_history'),
                (r'Current\s*Status[:\s]*', 'current_status'),
            ],
            'h6': [
                (r'General[:\s]*', 'general'),
                (r'Psychotic\s*Disorders?[:\s]*', 'psychotic_disorders'),
                (r'Major\s*Mood\s*Disorders?[:\s]*', 'major_mood_disorders'),
                (r'Other\s*(?:Major\s*)?Mental\s*Disorders?[:\s]*', 'other_mental_disorders'),
            ],
            'h7': [
                (r'Personality\s*Disorder\s*Features?[:\s]*', 'personality_disorder_features'),
                (r'Impact\s*on\s*Functioning[:\s]*', 'impact_on_functioning'),
            ],
            'h8': [
                (r'Victimi[sz]ation[/\s]*Trauma[:\s]*', 'victimization_trauma'),
                (r'Adverse\s*Childrearing\s*Experiences?[:\s]*', 'adverse_childrearing_experiences'),
            ],
            'h9': [
                (r'Violent\s*Attitudes?[:\s]*', 'violent_attitudes'),
                (r'Antisocial\s*Attitudes?[:\s]*', 'antisocial_attitudes'),
            ],
            'h10': [
                (r'Treatment\s*Response[:\s]*', 'treatment_response'),
                (r'Supervision\s*Response[:\s]*', 'supervision_response'),
            ],
            'c1': [
                (r'Insight\s*into\s*Mental\s*Health(?:\s*Diagnosis)?[:\s]*', 'insight_into_mental_health'),
                (r'Insight\s*into\s*Violence\s*Risk[:\s]*', 'insight_into_violence_risk'),
            ],
            'c2': [
                (r'Violent\s*Ideation[:\s]*', 'violent_ideation'),
                (r'Violent\s*Intent[:\s]*', 'violent_intent'),
            ],
            'r1': [
                (r'Hospital[:\s]*', 'hospital'),
                (r'Community[:\s]*', 'community'),
            ],
        }

        patterns = subsection_patterns.get(item_key, [])
        if not patterns:
            return {}

        result = {}

        # Build a combined pattern to split by any of the subsection headers
        all_patterns = [p[0] for p in patterns]
        combined = '|'.join(f'({p})' for p in all_patterns)

        # Find all matches and their positions
        matches = list(re.finditer(combined, content, re.IGNORECASE))

        if not matches:
            return {}

        for i, match in enumerate(matches):
            # Determine which pattern matched
            matched_text = match.group(0)
            field_suffix = None
            for pattern, suffix in patterns:
                if re.match(pattern, matched_text, re.IGNORECASE):
                    field_suffix = suffix
                    break

            if not field_suffix:
                continue

            # Get content from end of this match to start of next match (or end of string)
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)

            subsection_content = content[start:end].strip()
            if subsection_content:
                result[field_suffix] = subsection_content
                print(f"[HCR-20] Split {item_key}: {field_suffix} = {len(subsection_content)} chars")

        return result

    def _detect_and_parse_hcr20_document(self, file_path: str) -> dict:
        """
        Detect if a document is an HCR-20 report and parse its content.

        Returns a dict with H1-H10, C1-C5, R1-R5 content if it's an HCR-20 document,
        or None if it's not.
        """
        try:
            from docx import Document
            import re

            doc = Document(file_path)

            # Check if this is an HCR-20 document by looking for key indicators
            full_text = ' '.join([p.text for p in doc.paragraphs[:50]])

            is_hcr20 = False
            if 'HCR-20' in full_text or 'HCR20' in full_text:
                if any(x in full_text for x in ['Risk Assessment', 'Violence', 'Historical', 'Clinical']):
                    is_hcr20 = True

            if not is_hcr20:
                print(f"[HCR-20] Document does not appear to be an HCR-20 report")
                return None

            print(f"[HCR-20] Detected HCR-20 document - parsing...")

            # Parse the tables to extract H, C, R items
            hcr20_data = {
                'metadata': {},
                'items': {}
            }

            # Extract metadata from paragraphs
            for para in doc.paragraphs[:50]:
                text = para.text.strip()
                if 'NAME:' in text and 'name' not in hcr20_data['metadata']:
                    hcr20_data['metadata']['name'] = text.split('NAME:')[-1].strip()
                elif 'D.O.B:' in text or 'DOB:' in text:
                    hcr20_data['metadata']['dob'] = text.split(':')[-1].strip()
                elif 'AGE:' in text:
                    hcr20_data['metadata']['age'] = text.split('AGE:')[-1].strip()
                elif 'NHS NUMBER:' in text or 'NHS:' in text:
                    hcr20_data['metadata']['nhs_number'] = text.split(':')[-1].strip()
                elif 'ADDRESS:' in text:
                    hcr20_data['metadata']['address'] = text.split('ADDRESS:')[-1].strip()
                elif 'DATE OF ADMISSION:' in text or 'ADMISSION:' in text:
                    hcr20_data['metadata']['admission_date'] = text.split(':')[-1].strip()
                elif 'LEGAL STATUS:' in text:
                    hcr20_data['metadata']['legal_status'] = text.split('LEGAL STATUS:')[-1].strip()
                elif 'AUTHOR OF ORIGINAL' in text:
                    hcr20_data['metadata']['original_author'] = text.split(':')[-1].strip()
                elif 'AUTHOR OF UPDATE' in text:
                    hcr20_data['metadata']['update_author'] = text.split(':')[-1].strip()

            print(f"[HCR-20] Extracted metadata: {list(hcr20_data['metadata'].keys())}")

            # Parse tables for H, C, R items
            # Table structure: Each item has 4 rows (ITEM, Scoring, Presence, Relevance)
            # The content is in column 1 of row 1 (or spans rows 1-3)

            for table in doc.tables:
                for i, row in enumerate(table.rows):
                    cell0_text = row.cells[0].text.strip() if row.cells else ''

                    # Check if this is an ITEM row (H1, H2, C1, R1, etc.)
                    item_match = re.search(r'ITEM\s*(H\d+|C\d+|R\d+)', cell0_text, re.IGNORECASE)
                    if not item_match:
                        # Also check for items without "ITEM" prefix
                        item_match = re.search(r'^(H\d+|C\d+|R\d+)\b', cell0_text, re.IGNORECASE)

                    if item_match:
                        item_key = item_match.group(1).lower()  # e.g., 'h1', 'c1', 'r1'

                        # Get the title from column 1 of this row
                        title = row.cells[1].text.strip() if len(row.cells) > 1 else ''

                        # Get the content from the next rows (Scoring row has the main content)
                        content = ''
                        presence = ''
                        relevance = ''

                        # Look at the next 3 rows for Scoring, Presence, Relevance
                        for j in range(1, 4):
                            if i + j < len(table.rows):
                                next_row = table.rows[i + j]
                                next_cell0 = next_row.cells[0].text.strip().lower() if next_row.cells else ''
                                next_cell1 = next_row.cells[1].text.strip() if len(next_row.cells) > 1 else ''

                                if 'scoring' in next_cell0:
                                    content = next_cell1
                                elif 'presence' in next_cell0:
                                    # Extract presence level
                                    if 'present' in next_cell0.lower():
                                        if 'not' in next_cell0.lower():
                                            presence = 'Not Present'
                                        elif 'partial' in next_cell0.lower():
                                            presence = 'Partially Present'
                                        else:
                                            presence = 'Present'
                                elif 'relevance' in next_cell0:
                                    # Extract relevance level
                                    if 'high' in next_cell0.lower():
                                        relevance = 'High'
                                    elif 'moderate' in next_cell0.lower() or 'medium' in next_cell0.lower():
                                        relevance = 'Moderate'
                                    elif 'low' in next_cell0.lower():
                                        relevance = 'Low'

                        # Split content into subsections based on known header patterns
                        subsections = self._split_content_into_subsections(item_key, content)

                        hcr20_data['items'][item_key] = {
                            'title': title,
                            'content': content,
                            'subsections': subsections,
                            'presence': presence,
                            'relevance': relevance
                        }
                        print(f"[HCR-20] Parsed {item_key.upper()}: {len(content)} chars, {len(subsections)} subsections, presence={presence}, relevance={relevance}")

            # Also extract formulation and scenarios if present
            try:
                paragraphs = list(doc.paragraphs)
                formulation_start = None
                scenarios_start = None
                nature_start = None
                severity_start = None
                imminence_start = None
                frequency_start = None
                likelihood_start = None

                # Find section markers
                for idx, para in enumerate(paragraphs):
                    text = para.text.strip().lower()
                    if 'violence risk formulation' in text or 'risk formulation' in text:
                        formulation_start = idx
                    elif 'scenarios of potential risk' in text or text == 'scenarios':
                        scenarios_start = idx
                    elif 'nature' in text and ('violence' in text or 'committed' in text or idx == nature_start):
                        nature_start = idx
                    elif text.startswith('severity'):
                        severity_start = idx
                    elif text.startswith('imminence'):
                        imminence_start = idx
                    elif 'frequency' in text and 'duration' in text:
                        frequency_start = idx
                    elif text.startswith('likelihood'):
                        likelihood_start = idx

                # Extract formulation (from formulation header to scenarios)
                if formulation_start is not None:
                    end_idx = scenarios_start if scenarios_start else formulation_start + 50
                    formulation_text = []
                    for p in paragraphs[formulation_start+1:end_idx]:
                        t = p.text.strip()
                        if t and 'Scenarios' not in t:
                            formulation_text.append(t)
                        elif 'Scenarios' in t:
                            break
                    if formulation_text:
                        hcr20_data['formulation'] = '\n\n'.join(formulation_text)
                        print(f"[HCR-20] Parsed formulation: {len(hcr20_data['formulation'])} chars")

                # Extract scenarios
                def extract_section(start_idx, next_starts):
                    if start_idx is None:
                        return ''
                    # Find the end (next section or +30 paragraphs)
                    end_idx = start_idx + 30
                    for ns in next_starts:
                        if ns and ns > start_idx:
                            end_idx = min(end_idx, ns)
                            break
                    section_text = []
                    for p in paragraphs[start_idx+1:end_idx]:
                        t = p.text.strip()
                        if t:
                            # Stop if we hit another major section header
                            if any(h in t.lower() for h in ['severity', 'imminence', 'frequency', 'likelihood', 'proposed management', 'risk-enhancing']):
                                if t.lower() != paragraphs[start_idx].text.strip().lower():
                                    break
                            section_text.append(t)
                    return '\n\n'.join(section_text)

                hcr20_data['scenarios'] = {}

                if nature_start:
                    content = extract_section(nature_start, [severity_start, imminence_start, frequency_start, likelihood_start])
                    if content:
                        hcr20_data['scenarios']['nature'] = content
                        print(f"[HCR-20] Parsed scenario_nature: {len(content)} chars")

                if severity_start:
                    content = extract_section(severity_start, [imminence_start, frequency_start, likelihood_start])
                    if content:
                        hcr20_data['scenarios']['severity'] = content
                        print(f"[HCR-20] Parsed scenario_severity: {len(content)} chars")

                if imminence_start:
                    content = extract_section(imminence_start, [frequency_start, likelihood_start])
                    if content:
                        hcr20_data['scenarios']['imminence'] = content
                        print(f"[HCR-20] Parsed scenario_imminence: {len(content)} chars")

                if frequency_start:
                    content = extract_section(frequency_start, [likelihood_start])
                    if content:
                        hcr20_data['scenarios']['frequency'] = content
                        print(f"[HCR-20] Parsed scenario_frequency: {len(content)} chars")

                if likelihood_start:
                    content = extract_section(likelihood_start, [])
                    if content:
                        hcr20_data['scenarios']['likelihood'] = content
                        print(f"[HCR-20] Parsed scenario_likelihood: {len(content)} chars")

            except Exception as e:
                print(f"[HCR-20] Could not extract formulation/scenarios: {e}")
                import traceback
                traceback.print_exc()

            if hcr20_data['items']:
                print(f"[HCR-20] *** SUCCESS *** Parsed {len(hcr20_data['items'])} HCR-20 items")
                print(f"[HCR-20] Items: {list(hcr20_data['items'].keys())}")
                print(f"[HCR-20] Formulation: {'Yes' if hcr20_data.get('formulation') else 'No'}")
                print(f"[HCR-20] Scenarios: {list(hcr20_data.get('scenarios', {}).keys())}")
                print(f"[HCR-20] Returning parsed data to populate form...")
                return hcr20_data
            else:
                print(f"[HCR-20] No HCR-20 items found in tables")
                return None

        except Exception as e:
            print(f"[HCR-20] CRITICAL ERROR parsing HCR-20 document: {e}")
            import traceback
            traceback.print_exc()
            # Even if there was an error, if we have items, return them
            if 'hcr20_data' in locals() and hcr20_data.get('items'):
                print(f"[HCR-20] Returning partial data despite error...")
                return hcr20_data
            return None

    def _populate_from_hcr20_document(self, hcr20_data: dict, file_path: str):
        """
        Populate the HCR-20 form from a parsed HCR-20 document.
        Fills both cards AND popup widgets simultaneously.
        """
        from PySide6.QtWidgets import QMessageBox
        import os

        fname = os.path.basename(file_path)
        items_populated = 0

        print(f"[HCR-20] *** POPULATING FORM FROM HCR-20 DOCUMENT ***")
        print(f"[HCR-20] Available cards: {list(self.cards.keys())}")
        print(f"[HCR-20] Items to populate: {list(hcr20_data.get('items', {}).keys())}")
        print(f"[HCR-20] Has formulation: {bool(hcr20_data.get('formulation'))}")
        print(f"[HCR-20] Has scenarios: {list(hcr20_data.get('scenarios', {}).keys())}")

        # Populate patient details from metadata
        metadata = hcr20_data.get('metadata', {})
        if metadata:
            print(f"[HCR-20] Metadata to populate: {metadata}")

            if metadata.get('name') and hasattr(self, 'popup_patient_name'):
                self.popup_patient_name.setText(metadata['name'])
                print(f"[HCR-20] Populated patient name: {metadata['name']}")

            if metadata.get('nhs_number') and hasattr(self, 'popup_nhs'):
                self.popup_nhs.setText(metadata['nhs_number'])
                print(f"[HCR-20] Populated NHS number")

            if metadata.get('address') and hasattr(self, 'popup_address'):
                self.popup_address.setPlainText(metadata['address'])
                print(f"[HCR-20] Populated address")

            if metadata.get('legal_status') and hasattr(self, 'popup_legal_status'):
                self.popup_legal_status.setText(metadata['legal_status'])
                print(f"[HCR-20] Populated legal status")

            # Parse and set DOB (QDateEdit requires QDate)
            if metadata.get('dob') and hasattr(self, 'popup_dob'):
                try:
                    from dateutil import parser as date_parser
                    dob_date = date_parser.parse(metadata['dob'], dayfirst=True)
                    self.popup_dob.setDate(QDate(dob_date.year, dob_date.month, dob_date.day))
                    print(f"[HCR-20] Populated DOB: {metadata['dob']}")
                except Exception as e:
                    print(f"[HCR-20] Could not parse DOB '{metadata['dob']}': {e}")

            # Parse and set admission date
            if metadata.get('admission_date') and hasattr(self, 'popup_admission_date'):
                try:
                    from dateutil import parser as date_parser
                    adm_date = date_parser.parse(metadata['admission_date'], dayfirst=True)
                    self.popup_admission_date.setDate(QDate(adm_date.year, adm_date.month, adm_date.day))
                    print(f"[HCR-20] Populated admission date: {metadata['admission_date']}")
                except Exception as e:
                    print(f"[HCR-20] Could not parse admission date '{metadata['admission_date']}': {e}")

        # Populate each HCR-20 item (H1-H10, C1-C5, R1-R5)
        for item_key, item_data in hcr20_data.get('items', {}).items():
            content = item_data.get('content', '')
            presence = item_data.get('presence', '')
            relevance = item_data.get('relevance', '')

            if not content:
                continue

            # 1. Populate the card editor
            if item_key in self.cards:
                card = self.cards[item_key]
                if hasattr(card, 'editor'):
                    card.editor.setPlainText(content)
                    items_populated += 1
                    print(f"[HCR-20] Populated card {item_key.upper()} with {len(content)} chars")

            # 2. Populate the popup subsection fields
            subsections = item_data.get('subsections', {})

            if subsections:
                # Item has subsections - populate each one
                for field_suffix, subsection_content in subsections.items():
                    # The field names are popup_{key}_{sanitized_label}
                    # We need to match the field suffix to the actual attribute name
                    # Field suffix from parsing: e.g., 'child_(aged_12_and_under)'
                    # Actual attribute: e.g., 'popup_h1_child_(aged_12_and_under)'
                    attr_name = f"popup_{item_key}_{field_suffix}"

                    if hasattr(self, attr_name):
                        widget = getattr(self, attr_name)
                        if hasattr(widget, 'setPlainText'):
                            widget.setPlainText(subsection_content)
                            print(f"[HCR-20] Populated {attr_name} with {len(subsection_content)} chars")
                    else:
                        # Try to find a matching field by checking actual attributes
                        # (field names may have slight variations)
                        for actual_attr in dir(self):
                            if actual_attr.startswith(f"popup_{item_key}_"):
                                # Check if this is a text field (not presence/relevance/imported)
                                if any(x in actual_attr for x in ['presence', 'relevance', 'imported', 'checkboxes', 'group']):
                                    continue
                                # Check if the field suffix is similar
                                suffix_clean = field_suffix.lower().replace('(', '').replace(')', '').replace('+', '')
                                attr_clean = actual_attr.lower().replace('(', '').replace(')', '').replace('+', '')
                                if suffix_clean in attr_clean or attr_clean.endswith(suffix_clean):
                                    widget = getattr(self, actual_attr, None)
                                    if widget and hasattr(widget, 'setPlainText'):
                                        widget.setPlainText(subsection_content)
                                        print(f"[HCR-20] Populated {actual_attr} (matched from {field_suffix}) with {len(subsection_content)} chars")
                                        break
            else:
                # No subsections - use popup_{key}_evidence for items without subsections
                evidence_attr = f"popup_{item_key}_evidence"
                if hasattr(self, evidence_attr):
                    getattr(self, evidence_attr).setPlainText(content)
                    print(f"[HCR-20] Populated {item_key} popup evidence with {len(content)} chars")
                else:
                    # Fallback: find first text field for this item
                    for attr_name in dir(self):
                        if attr_name.startswith(f"popup_{item_key}_") and not any(x in attr_name for x in ['presence', 'relevance', 'imported', 'checkboxes', 'group']):
                            widget = getattr(self, attr_name, None)
                            if widget and hasattr(widget, 'setPlainText'):
                                widget.setPlainText(content)
                                print(f"[HCR-20] Populated {attr_name} (fallback) with {len(content)} chars")
                                break

            # 3. Set presence radio button
            if presence:
                presence_map = {
                    'present': f"popup_{item_key}_presence_yes",
                    'partially present': f"popup_{item_key}_presence_partial",
                    'partial': f"popup_{item_key}_presence_partial",
                    'not present': f"popup_{item_key}_presence_no",
                    'omit': f"popup_{item_key}_presence_omit",
                }
                for pres_key, attr_name in presence_map.items():
                    if pres_key in presence.lower():
                        if hasattr(self, attr_name):
                            getattr(self, attr_name).setChecked(True)
                            print(f"[HCR-20] Set {item_key} presence to {pres_key}")
                            break

            # 4. Set relevance radio button
            if relevance:
                relevance_map = {
                    'high': f"popup_{item_key}_relevance_high",
                    'moderate': f"popup_{item_key}_relevance_mod",
                    'low': f"popup_{item_key}_relevance_low",
                }
                for rel_key, attr_name in relevance_map.items():
                    if rel_key in relevance.lower():
                        if hasattr(self, attr_name):
                            getattr(self, attr_name).setChecked(True)
                            print(f"[HCR-20] Set {item_key} relevance to {rel_key}")
                            break

        # Populate formulation
        if 'formulation' in hcr20_data:
            formulation_content = hcr20_data['formulation']
            # Popup text widget
            if hasattr(self, 'popup_formulation_text'):
                self.popup_formulation_text.setPlainText(formulation_content)
                print(f"[HCR-20] Populated formulation popup with {len(formulation_content)} chars")
                items_populated += 1
            # Card
            if 'formulation' in self.cards:
                card = self.cards['formulation']
                if hasattr(card, 'editor'):
                    card.editor.setPlainText(formulation_content)

        # Populate scenarios
        scenarios = hcr20_data.get('scenarios', {})
        scenario_map = {
            'nature': 'scenario_nature',
            'severity': 'scenario_severity',
            'imminence': 'scenario_imminence',
            'frequency': 'scenario_frequency',
            'likelihood': 'scenario_likelihood',
        }

        for scenario_key, card_key in scenario_map.items():
            content = scenarios.get(scenario_key, '')
            if content:
                # Popup text widget (named popup_{card_key}_text)
                popup_attr = f"popup_{card_key}_text"
                if hasattr(self, popup_attr):
                    getattr(self, popup_attr).setPlainText(content)
                    print(f"[HCR-20] Populated {card_key} popup with {len(content)} chars")
                    items_populated += 1
                # Card
                if card_key in self.cards:
                    card = self.cards[card_key]
                    if hasattr(card, 'editor'):
                        card.editor.setPlainText(content)

        # Update all popup previews to reflect the new content
        for key in list(self.cards.keys()):
            try:
                self._update_preview(key)
            except Exception:
                pass  # Silent fail for preview updates

        # Show success message
        msg = f"HCR-20 Report Imported Successfully!\n\n"
        msg += f"Source: {fname}\n\n"
        msg += f"Populated {items_populated} sections.\n\n"

        if metadata.get('name'):
            msg += f"Patient: {metadata['name']}\n"

        items_count = len(hcr20_data.get('items', {}))
        msg += f"✓ {items_count} HCR-20 items (H1-H10, C1-C5, R1-R5)\n"

        if hcr20_data.get('formulation'):
            msg += "✓ Violence Risk Formulation\n"
        if scenarios:
            msg += f"✓ {len(scenarios)} Scenario sections\n"

        msg += "\nThe form has been pre-filled with the imported data.\n"
        msg += "Review and edit as needed."

        QMessageBox.information(self, "HCR-20 Import Complete", msg)

    # ================================================================
    # FORENSIC DATA FOR H1/H2 - EXACT COPY FROM TRIBUNAL REPORT SECTION 5
    # ================================================================

    def set_hcr_forensic_data(self, key: str, notes: list):
        """Set forensic data for H1 or H2 using the same approach as tribunal report section 5.

        For H1 (Violence): Analyzes Physical Aggression, Verbal Aggression (high severity)
        For H2 (Antisocial): Analyzes Property Damage, Sexual Behaviour
        """
        from risk_overview_panel import analyze_notes_for_risk
        from datetime import datetime
        import re
        import html

        print(f"[HCR-20] set_hcr_forensic_data called for {key} with {len(notes) if notes else 0} notes")

        # Define which risk categories to use for each item
        if key == 'h1':
            RELEVANT_CATEGORIES = ["Physical Aggression", "Verbal Aggression"]
        elif key == 'h2':
            RELEVANT_CATEGORIES = ["Property Damage", "Sexual Behaviour"]
        else:
            RELEVANT_CATEGORIES = ["Physical Aggression", "Property Damage"]

        all_incidents = []

        # Run risk analysis on notes
        if notes:
            results = analyze_notes_for_risk(notes)
            for cat_name in RELEVANT_CATEGORIES:
                cat_data = results.get("categories", {}).get(cat_name, {})
                for incident in cat_data.get("incidents", []):
                    all_incidents.append({
                        "date": incident.get("date"),
                        "text": incident.get("full_text", ""),
                        "matched": incident.get("matched", ""),
                        "subcategory": incident.get("subcategory", ""),
                        "severity": incident.get("severity", "medium"),
                        "category": cat_name,
                        "source": "notes",
                    })

        if not all_incidents:
            print(f"[HCR-20] No incidents found for {key}")
            return

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                return datetime.min
            return d

        sorted_incidents = sorted(all_incidents, key=get_sort_date, reverse=True)

        # Store incidents for this key
        setattr(self, f"_{key}_all_incidents", sorted_incidents)
        setattr(self, f"_{key}_current_filter", None)

        # Category colors
        cat_colors = {
            "Physical Aggression": "#b71c1c",
            "Verbal Aggression": "#9E9E9E",
            "Property Damage": "#e53935",
            "Sexual Behaviour": "#673ab7",
        }

        severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        setattr(self, f"_{key}_cat_colors", cat_colors)
        setattr(self, f"_{key}_severity_colors", severity_colors)

        # Get the imported entries layout for this key
        entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
        if not entries_layout:
            print(f"[HCR-20] No imported entries layout found for {key}")
            return

        # Clear existing content
        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Collect unique labels (category or category:subcategory)
        labels = {}
        for inc in sorted_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat:
                label = f"{cat}: {subcat}"
            else:
                label = cat
            if label not in labels:
                labels[label] = cat_colors.get(cat, "#666666")

        # Create filter panel
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(26)  # Reduced by 10% (was 29)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollBar:horizontal { height: 6px; background: #f0f0f0; border-radius: 3px; }
            QScrollBar::handle:horizontal { background: #c0c0c0; border-radius: 3px; min-width: 20px; }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 4px 10px;
                    font-size: 12px;
                    font-weight: 600;
                }}
                QPushButton:hover {{ background: {color}dd; }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text, k=key: self._apply_hcr_forensic_filter(k, lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row
        filter_status_widget = QWidget()
        filter_status_widget.setStyleSheet("background: transparent;")
        filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        filter_label = QLabel("Filtered by: ")
        filter_label.setStyleSheet("font-size: 13px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton { background: #ef4444; color: white; border: none; border-radius: 4px; padding: 3px 8px; font-size: 12px; font-weight: 600; }
            QPushButton:hover { background: #dc2626; }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(lambda _, k=key: self._remove_hcr_forensic_filter(k))
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(filter_status_widget)
        entries_layout.addWidget(filter_container)

        # Store filter widgets
        setattr(self, f"_{key}_filter_status_widget", filter_status_widget)
        setattr(self, f"_{key}_filter_label", filter_label)

        # Container for incident entries
        incidents_container = QWidget()
        incidents_container.setStyleSheet("background: transparent;")
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(6)
        entries_layout.addWidget(incidents_container)

        setattr(self, f"_{key}_incidents_layout", incidents_layout)

        # Render incidents
        self._render_hcr_forensic_incidents(key, sorted_incidents)

        # Show the imported section
        imported_section = getattr(self, f"popup_{key}_imported_section", None)
        if imported_section:
            imported_section.setVisible(True)

        print(f"[HCR-20] Displayed {len(sorted_incidents)} incidents for {key}")

    def _render_hcr_forensic_incidents(self, key: str, incidents: list):
        """Render forensic incidents for H1/H2 - exact copy from tribunal report."""
        import re
        import html

        incidents_layout = getattr(self, f"_{key}_incidents_layout", None)
        if not incidents_layout:
            return

        cat_colors = getattr(self, f"_{key}_cat_colors", {})
        severity_colors = getattr(self, f"_{key}_severity_colors", {})

        # Clear existing
        while incidents_layout.count():
            child = incidents_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        for incident in incidents:
            date = incident["date"]
            cat_name = incident["category"]
            text = incident["text"]
            matched = incident["matched"]
            subcat_name = incident["subcategory"]
            severity = incident["severity"]

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Get colors
            cat_color = cat_colors.get(cat_name, "#666666")
            sev_color = severity_colors.get(severity, "#666666")

            # Create HTML with highlighted matched text
            escaped_text = html.escape(text)
            if matched:
                escaped_matched = html.escape(matched)
                try:
                    pattern = re.compile(re.escape(escaped_matched), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_matched}</span>',
                        escaped_text
                    )
                except:
                    highlighted_html = escaped_text
            else:
                highlighted_html = escaped_text

            full_html = f'''
            <html>
            <body style="font-family: -apple-system, BlinkMacSystemFont, sans-serif; font-size: 14px; color: #333; margin: 0; padding: 0;">
            {highlighted_html}
            </body>
            </html>
            '''

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("hcrForensicEntry")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#hcrForensicEntry {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid #e5e7eb;
                    border-left: 4px solid {cat_color};
                    border-radius: 6px;
                    padding: 2px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(6, 4, 6, 4)
            entry_layout.setSpacing(3)

            # Header row with badges and date
            header_row = QHBoxLayout()
            header_row.setSpacing(6)

            # Category badge
            badge_text = f"{cat_name}: {subcat_name}" if subcat_name else cat_name
            cat_badge = QLabel(badge_text)
            cat_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 10px;
                    font-weight: 600;
                    color: white;
                    background: {cat_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 5px;
                }}
            """)
            header_row.addWidget(cat_badge)

            # Severity badge
            sev_badge = QLabel(severity.upper())
            sev_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 9px;
                    font-weight: 700;
                    color: white;
                    background: {sev_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 4px;
                }}
            """)
            header_row.addWidget(sev_badge)

            # Date label
            date_label = QLabel(f"📅 {date_str}")
            date_label.setStyleSheet("font-size: 12px; font-weight: 500; color: #6b7280; background: transparent; border: none;")
            header_row.addWidget(date_label)

            header_row.addStretch()
            entry_layout.addLayout(header_row)

            # Text content
            body_text = QTextEdit()
            body_text.setReadOnly(True)
            body_text.setHtml(full_html)
            body_text.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
            body_text.setMinimumHeight(36)  # Reduced by 10% (was 40)
            body_text.setMaximumHeight(72)  # Reduced by 10% (was 80)
            body_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
            body_text.setStyleSheet("""
                QTextEdit {
                    background: #f9fafb;
                    border: none;
                    border-radius: 4px;
                    padding: 4px;
                    font-size: 14px;
                    color: #374151;
                }
            """)
            entry_layout.addWidget(body_text)

            incidents_layout.addWidget(entry_frame)

    def _apply_hcr_forensic_filter(self, key: str, label: str):
        """Apply filter for H1/H2 forensic data."""
        setattr(self, f"_{key}_current_filter", label)

        filter_label = getattr(self, f"_{key}_filter_label", None)
        filter_status_widget = getattr(self, f"_{key}_filter_status_widget", None)

        if filter_label:
            filter_label.setText(f"Filtered by: {label}")
        if filter_status_widget:
            filter_status_widget.setVisible(True)

        # Filter incidents
        all_incidents = getattr(self, f"_{key}_all_incidents", [])
        filtered = []
        for inc in all_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat:
                inc_label = f"{cat}: {subcat}"
            else:
                inc_label = cat
            if inc_label == label:
                filtered.append(inc)

        self._render_hcr_forensic_incidents(key, filtered)

    def _remove_hcr_forensic_filter(self, key: str):
        """Remove filter for H1/H2 forensic data."""
        setattr(self, f"_{key}_current_filter", None)

        filter_status_widget = getattr(self, f"_{key}_filter_status_widget", None)
        if filter_status_widget:
            filter_status_widget.setVisible(False)

        all_incidents = getattr(self, f"_{key}_all_incidents", [])
        self._render_hcr_forensic_incidents(key, all_incidents)

    def _set_hcr_forensic_data_cached(self, key: str, cached_results: dict):
        """Set forensic data using pre-computed risk analysis results (faster)."""
        from datetime import datetime
        import re
        import html

        print(f"[HCR-20] _set_hcr_forensic_data_cached called for {key}")

        # Define which risk categories to use for each item
        if key == 'h1':
            RELEVANT_CATEGORIES = ["Physical Aggression", "Verbal Aggression"]
        elif key == 'h2':
            RELEVANT_CATEGORIES = ["Property Damage", "Sexual Behaviour"]
        else:
            RELEVANT_CATEGORIES = ["Physical Aggression", "Property Damage"]

        all_incidents = []

        # Use cached results instead of running analyze_notes_for_risk again
        for cat_name in RELEVANT_CATEGORIES:
            cat_data = cached_results.get("categories", {}).get(cat_name, {})
            for incident in cat_data.get("incidents", []):
                all_incidents.append({
                    "date": incident.get("date"),
                    "text": incident.get("full_text", ""),
                    "matched": incident.get("matched", ""),
                    "subcategory": incident.get("subcategory", ""),
                    "severity": incident.get("severity", "medium"),
                    "category": cat_name,
                    "source": "notes",
                })

        if not all_incidents:
            print(f"[HCR-20] No incidents found for {key}")
            return

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                return datetime.min
            return d

        sorted_incidents = sorted(all_incidents, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep only the most relevant entry per date
        seen_dates = {}
        deduplicated_incidents = []
        for inc in sorted_incidents:
            d = inc.get("date")
            if d is None:
                date_key = ""
            elif hasattr(d, "strftime"):
                date_key = d.strftime("%Y-%m-%d")
            else:
                date_key = str(d)

            severity_weights = {"high": 300, "medium": 200, "low": 100}
            text_len = len(inc.get('text', ''))
            relevance_score = severity_weights.get(inc.get('severity', 'medium'), 200) + text_len

            if date_key not in seen_dates:
                seen_dates[date_key] = (inc, relevance_score)
                deduplicated_incidents.append(inc)
            else:
                existing_inc, existing_score = seen_dates[date_key]
                if relevance_score > existing_score:
                    deduplicated_incidents.remove(existing_inc)
                    deduplicated_incidents.append(inc)
                    seen_dates[date_key] = (inc, relevance_score)

        sorted_incidents = sorted(deduplicated_incidents, key=get_sort_date, reverse=True)

        # Store incidents for this key
        setattr(self, f"_{key}_all_incidents", sorted_incidents)
        setattr(self, f"_{key}_current_filter", None)

        # Category colors
        cat_colors = {
            "Physical Aggression": "#b71c1c",
            "Verbal Aggression": "#9E9E9E",
            "Property Damage": "#e53935",
            "Sexual Behaviour": "#673ab7",
        }

        severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        setattr(self, f"_{key}_cat_colors", cat_colors)
        setattr(self, f"_{key}_severity_colors", severity_colors)

        # Get the imported entries layout for this key
        entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
        if not entries_layout:
            print(f"[HCR-20] No imported entries layout found for {key}")
            return

        # Clear existing content
        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Collect unique labels
        labels = {}
        for inc in sorted_incidents:
            cat = inc["category"]
            subcat = inc.get("subcategory", "")
            if subcat:
                label = f"{cat}: {subcat}"
            else:
                label = cat
            if label not in labels:
                labels[label] = cat_colors.get(cat, "#666666")

        # Create filter panel
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(26)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollBar:horizontal { height: 6px; background: #f0f0f0; border-radius: 3px; }
            QScrollBar::handle:horizontal { background: #c0c0c0; border-radius: 3px; min-width: 20px; }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 4px 10px;
                    font-size: 12px;
                    font-weight: 600;
                }}
                QPushButton:hover {{ background: {color}dd; }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text, k=key: self._apply_hcr_forensic_filter(k, lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row
        filter_status_widget = QWidget()
        filter_status_widget.setStyleSheet("background: transparent;")
        filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        filter_label = QLabel("Filtered by: ")
        filter_label.setStyleSheet("font-size: 13px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton { background: #ef4444; color: white; border: none; border-radius: 4px; padding: 3px 8px; font-size: 12px; font-weight: 600; }
            QPushButton:hover { background: #dc2626; }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(lambda _, k=key: self._remove_hcr_forensic_filter(k))
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(filter_status_widget)
        entries_layout.addWidget(filter_container)

        # Store filter widgets
        setattr(self, f"_{key}_filter_status_widget", filter_status_widget)
        setattr(self, f"_{key}_filter_label", filter_label)

        # Container for incident entries
        incidents_container = QWidget()
        incidents_container.setStyleSheet("background: transparent;")
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(6)
        entries_layout.addWidget(incidents_container)

        setattr(self, f"_{key}_incidents_layout", incidents_layout)

        # Render incidents
        self._render_hcr_forensic_incidents(key, sorted_incidents)

        # Show the imported section
        imported_section = getattr(self, f"popup_{key}_imported_section", None)
        if imported_section:
            imported_section.setVisible(True)

        print(f"[HCR-20] Displayed {len(sorted_incidents)} incidents for {key}")

    # ================================================================
    # SUBSTANCE DATA FOR H5 - USING SAME APPROACH AS GPR SECTION 8
    # ================================================================

    def set_hcr_substance_data(self, notes: list):
        """Set substance data for H5 using the same approach as GPR section 8.

        Uses analyze_notes_for_risk to find "Substance Misuse" category incidents.
        """
        from risk_overview_panel import analyze_notes_for_risk
        from datetime import datetime
        import re
        import html

        key = 'h5'
        print(f"[HCR-20] set_hcr_substance_data called with {len(notes) if notes else 0} notes")

        all_incidents = []

        # Run risk analysis on notes - use "Substance Misuse" category
        if notes:
            results = analyze_notes_for_risk(notes)
            substance_data = results.get("categories", {}).get("Substance Misuse", {})
            for incident in substance_data.get("incidents", []):
                all_incidents.append({
                    "date": incident.get("date"),
                    "text": incident.get("full_text", ""),
                    "matched": incident.get("matched", ""),
                    "subcategory": incident.get("subcategory", ""),
                    "severity": incident.get("severity", "medium"),
                    "category": "Substance Misuse",
                    "source": "notes",
                })

        if not all_incidents:
            print(f"[HCR-20] No substance incidents found for H5")
            return

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                return datetime.min
            return d

        sorted_incidents = sorted(all_incidents, key=get_sort_date, reverse=True)

        # Store incidents for this key
        setattr(self, f"_{key}_all_incidents", sorted_incidents)
        setattr(self, f"_{key}_current_filter", None)

        # Subcategory colors (from GPR section 8)
        subcat_colors = {
            "Positive Drug Test": "#6a1b9a",
            "Smelling of Substances": "#7b1fa2",
            "Appeared Intoxicated": "#8e24aa",
            "Admitted Substance Use": "#9c27b0",
            "Found with Substances": "#ab47bc",
        }

        severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        setattr(self, f"_{key}_subcat_colors", subcat_colors)
        setattr(self, f"_{key}_severity_colors", severity_colors)

        # Get the imported entries layout for this key
        entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
        if not entries_layout:
            print(f"[HCR-20] No imported entries layout found for {key}")
            return

        # Clear existing content
        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Collect unique labels
        labels = {}
        for inc in sorted_incidents:
            subcat = inc.get("subcategory", "")
            label = f"Substance Misuse: {subcat}" if subcat else "Substance Misuse"
            if label not in labels:
                labels[label] = subcat_colors.get(subcat, "#9c27b0")

        # Create filter panel
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(26)  # Reduced by 10%
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollBar:horizontal { height: 6px; background: #f0f0f0; border-radius: 3px; }
            QScrollBar::handle:horizontal { background: #c0c0c0; border-radius: 3px; min-width: 20px; }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 4px 10px;
                    font-size: 12px;
                    font-weight: 600;
                }}
                QPushButton:hover {{ background: {color}dd; }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text: self._apply_hcr_substance_filter(lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row
        filter_status_widget = QWidget()
        filter_status_widget.setStyleSheet("background: transparent;")
        filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        filter_label = QLabel("Filtered by: ")
        filter_label.setStyleSheet("font-size: 13px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton { background: #ef4444; color: white; border: none; border-radius: 4px; padding: 3px 8px; font-size: 12px; font-weight: 600; }
            QPushButton:hover { background: #dc2626; }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(self._remove_hcr_substance_filter)
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(filter_status_widget)
        entries_layout.addWidget(filter_container)

        # Store filter widgets
        setattr(self, f"_{key}_filter_status_widget", filter_status_widget)
        setattr(self, f"_{key}_filter_label", filter_label)

        # Container for incident entries
        incidents_container = QWidget()
        incidents_container.setStyleSheet("background: transparent;")
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(6)
        entries_layout.addWidget(incidents_container)

        setattr(self, f"_{key}_incidents_layout", incidents_layout)

        # Render incidents
        self._render_hcr_substance_incidents(sorted_incidents)

        # Show the imported section
        imported_section = getattr(self, f"popup_{key}_imported_section", None)
        if imported_section:
            imported_section.setVisible(True)
            if hasattr(imported_section, '_is_collapsed') and imported_section._is_collapsed:
                imported_section._toggle_collapse()

        print(f"[HCR-20] Displayed {len(sorted_incidents)} substance incidents for H5")

    def _render_hcr_substance_incidents(self, incidents: list):
        """Render substance incidents for H5 - matching H1 toggle/expand style."""
        import re
        import html

        key = 'h5'
        incidents_layout = getattr(self, f"_{key}_incidents_layout", None)
        if not incidents_layout:
            return

        subcat_colors = getattr(self, f"_{key}_subcat_colors", {})

        # Get or create checkboxes list
        checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        # Clear existing
        while incidents_layout.count():
            child = incidents_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        accent_color = "#9c27b0"  # Purple for substance
        bg_color = "rgba(243, 229, 245, 0.95)"
        border_color = "rgba(156, 39, 176, 0.4)"

        for incident in incidents:
            date = incident["date"]
            text = incident["text"]
            matched = incident.get("matched", "")
            subcat_name = incident.get("subcategory", "")

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Create HTML with highlighted matched text
            escaped_text = html.escape(text)
            highlighted_html = escaped_text
            if matched:
                try:
                    escaped_matched = html.escape(matched)
                    pattern = re.compile(re.escape(escaped_matched), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_matched}</span>',
                        highlighted_html
                    )
                except:
                    pass
            highlighted_html = highlighted_html.replace('\n', '<br>')

            # Create entry frame with colored left border (matching H1 style)
            entry_frame = QFrame()
            entry_frame.setObjectName("substanceEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            subcat_color = subcat_colors.get(subcat_name, accent_color)
            entry_frame.setStyleSheet(f"""
                QFrame#substanceEntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {subcat_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row with toggle, date, badges, and checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button
            toggle_btn = QPushButton("\u25b8")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(f"{date_str}")
            date_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 16px;
                    font-weight: 600;
                    color: {accent_color};
                    background: transparent;
                    border: none;
                }}
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Subcategory badge
            if subcat_name:
                subcat_badge = QLabel(subcat_name)
                subcat_badge.setStyleSheet(f"""
                    QLabel {{
                        font-size: 11px;
                        font-weight: 600;
                        color: white;
                        background: {subcat_color}cc;
                        border: none;
                        border-radius: 3px;
                        padding: 2px 6px;
                    }}
                """)
                header_row.addWidget(subcat_badge)

            header_row.addStretch()

            # Checkbox
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body text (hidden by default, matching H1)
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"""
                QTextEdit {{
                    font-size: 15px;
                    color: #333;
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }}
            """)
            body_text.setVisible(False)

            # Calculate height based on content
            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))

            # Toggle function
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("\u25b8")
                    else:
                        body.setVisible(True)
                        btn.setText("\u25be")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout.addWidget(body_text)
            incidents_layout.addWidget(entry_frame)

        # Store checkboxes reference
        setattr(self, f"popup_{key}_imported_checkboxes", checkboxes)

    def _apply_hcr_substance_filter(self, label: str):
        """Apply filter for H5 substance data."""
        key = 'h5'
        setattr(self, f"_{key}_current_filter", label)

        filter_label = getattr(self, f"_{key}_filter_label", None)
        filter_status_widget = getattr(self, f"_{key}_filter_status_widget", None)

        if filter_label:
            filter_label.setText(f"Filtered by: {label}")
        if filter_status_widget:
            filter_status_widget.setVisible(True)

        # Filter incidents
        all_incidents = getattr(self, f"_{key}_all_incidents", [])
        filtered = []
        for inc in all_incidents:
            subcat = inc.get("subcategory", "")
            inc_label = f"Substance Misuse: {subcat}" if subcat else "Substance Misuse"
            if inc_label == label:
                filtered.append(inc)

        self._render_hcr_substance_incidents(filtered)

    def _remove_hcr_substance_filter(self):
        """Remove filter for H5 substance data."""
        key = 'h5'
        setattr(self, f"_{key}_current_filter", None)

        filter_status_widget = getattr(self, f"_{key}_filter_status_widget", None)
        if filter_status_widget:
            filter_status_widget.setVisible(False)

        all_incidents = getattr(self, f"_{key}_all_incidents", [])
        self._render_hcr_substance_incidents(all_incidents)

    def _set_hcr_substance_data_cached(self, cached_results: dict):
        """Set substance data using pre-computed risk analysis results (faster)."""
        from datetime import datetime
        import re
        import html

        key = 'h5'
        print(f"[HCR-20] _set_hcr_substance_data_cached called")

        all_incidents = []

        # Use cached results instead of running analyze_notes_for_risk again
        substance_data = cached_results.get("categories", {}).get("Substance Misuse", {})
        for incident in substance_data.get("incidents", []):
            all_incidents.append({
                "date": incident.get("date"),
                "text": incident.get("full_text", ""),
                "matched": incident.get("matched", ""),
                "subcategory": incident.get("subcategory", ""),
                "severity": incident.get("severity", "medium"),
                "category": "Substance Misuse",
                "source": "notes",
            })

        if not all_incidents:
            print(f"[HCR-20] No substance incidents found for H5")
            return

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                return datetime.min
            return d

        sorted_incidents = sorted(all_incidents, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep only the most relevant entry per date
        seen_dates = {}
        deduplicated_incidents = []
        for inc in sorted_incidents:
            d = inc.get("date")
            if d is None:
                date_key = ""
            elif hasattr(d, "strftime"):
                date_key = d.strftime("%Y-%m-%d")
            else:
                date_key = str(d)

            severity_weights = {"high": 300, "medium": 200, "low": 100}
            text_len = len(inc.get('text', ''))
            relevance_score = severity_weights.get(inc.get('severity', 'medium'), 200) + text_len

            if date_key not in seen_dates:
                seen_dates[date_key] = (inc, relevance_score)
                deduplicated_incidents.append(inc)
            else:
                existing_inc, existing_score = seen_dates[date_key]
                if relevance_score > existing_score:
                    deduplicated_incidents.remove(existing_inc)
                    deduplicated_incidents.append(inc)
                    seen_dates[date_key] = (inc, relevance_score)

        sorted_incidents = sorted(deduplicated_incidents, key=get_sort_date, reverse=True)

        # Store incidents for this key
        setattr(self, f"_{key}_all_incidents", sorted_incidents)
        setattr(self, f"_{key}_current_filter", None)

        # Subcategory colors
        subcat_colors = {
            "Positive Drug Test": "#6a1b9a",
            "Smelling of Substances": "#7b1fa2",
            "Appeared Intoxicated": "#8e24aa",
            "Admitted Substance Use": "#9c27b0",
            "Found with Substances": "#ab47bc",
        }

        severity_colors = {
            "high": "#dc2626",
            "medium": "#f59e0b",
            "low": "#22c55e",
        }

        setattr(self, f"_{key}_subcat_colors", subcat_colors)
        setattr(self, f"_{key}_severity_colors", severity_colors)

        # Get the imported entries layout for this key
        entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
        if not entries_layout:
            print(f"[HCR-20] No imported entries layout found for {key}")
            return

        # Clear existing content
        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Collect unique labels
        labels = {}
        for inc in sorted_incidents:
            subcat = inc.get("subcategory", "")
            label = f"Substance Misuse: {subcat}" if subcat else "Substance Misuse"
            if label not in labels:
                labels[label] = subcat_colors.get(subcat, "#9c27b0")

        # Create filter panel
        filter_container = QWidget()
        filter_container.setStyleSheet("background: transparent;")
        filter_layout = QVBoxLayout(filter_container)
        filter_layout.setContentsMargins(0, 0, 0, 8)
        filter_layout.setSpacing(6)

        # Horizontal scroll for labels
        label_scroll = QScrollArea()
        label_scroll.setWidgetResizable(True)
        label_scroll.setFixedHeight(26)
        label_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        label_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        label_scroll.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollBar:horizontal { height: 6px; background: #f0f0f0; border-radius: 3px; }
            QScrollBar::handle:horizontal { background: #c0c0c0; border-radius: 3px; min-width: 20px; }
        """)

        label_widget = QWidget()
        label_widget.setStyleSheet("background: transparent;")
        label_row = QHBoxLayout(label_widget)
        label_row.setContentsMargins(0, 0, 0, 0)
        label_row.setSpacing(6)

        for label_text, color in sorted(labels.items()):
            btn = QPushButton(label_text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 4px;
                    padding: 4px 10px;
                    font-size: 12px;
                    font-weight: 600;
                }}
                QPushButton:hover {{ background: {color}dd; }}
            """)
            btn.setCursor(Qt.CursorShape.PointingHandCursor)
            btn.clicked.connect(lambda checked, lbl=label_text: self._apply_hcr_substance_filter(lbl))
            label_row.addWidget(btn)

        label_row.addStretch()
        label_scroll.setWidget(label_widget)
        filter_layout.addWidget(label_scroll)

        # Filter status row
        filter_status_widget = QWidget()
        filter_status_widget.setStyleSheet("background: transparent;")
        filter_status_widget.setVisible(False)
        filter_status_layout = QHBoxLayout(filter_status_widget)
        filter_status_layout.setContentsMargins(0, 0, 0, 0)
        filter_status_layout.setSpacing(8)

        filter_label = QLabel("Filtered by: ")
        filter_label.setStyleSheet("font-size: 13px; color: #374151; font-weight: 500;")
        filter_status_layout.addWidget(filter_label)

        remove_filter_btn = QPushButton("✕ Remove filter")
        remove_filter_btn.setStyleSheet("""
            QPushButton { background: #ef4444; color: white; border: none; border-radius: 4px; padding: 3px 8px; font-size: 12px; font-weight: 600; }
            QPushButton:hover { background: #dc2626; }
        """)
        remove_filter_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        remove_filter_btn.clicked.connect(self._remove_hcr_substance_filter)
        filter_status_layout.addWidget(remove_filter_btn)
        filter_status_layout.addStretch()

        filter_layout.addWidget(filter_status_widget)
        entries_layout.addWidget(filter_container)

        # Store filter widgets
        setattr(self, f"_{key}_filter_status_widget", filter_status_widget)
        setattr(self, f"_{key}_filter_label", filter_label)

        # Container for incident entries
        incidents_container = QWidget()
        incidents_container.setStyleSheet("background: transparent;")
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(6)
        entries_layout.addWidget(incidents_container)

        setattr(self, f"_{key}_incidents_layout", incidents_layout)

        # Render incidents
        self._render_hcr_substance_incidents(sorted_incidents)

        # Show the imported section
        imported_section = getattr(self, f"popup_{key}_imported_section", None)
        if imported_section:
            imported_section.setVisible(True)
            if hasattr(imported_section, '_is_collapsed') and imported_section._is_collapsed:
                imported_section._toggle_collapse()

        print(f"[HCR-20] Displayed {len(sorted_incidents)} substance incidents for H5")

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file for HCR-20 extraction."""
        import os

        # Check if this is an HCR-20 document
        ext = os.path.splitext(file_path)[1].lower()
        if ext in ['.docx', '.doc']:
            hcr20_data = self._detect_and_parse_hcr20_document(file_path)
            if hcr20_data:
                self._populate_from_hcr20_document(hcr20_data, file_path)
                return

        # Import as clinical notes
        all_notes = []
        try:
            notes = self._ingest_file(file_path)
            if notes:
                all_notes.extend(notes)
                print(f"[HCR-20] Imported {len(notes)} notes from {os.path.basename(file_path)}")
        except Exception as e:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(
                self, "Import Failed",
                f"Could not import notes:\n\n{os.path.basename(file_path)}: {str(e)}"
            )
            return

        if not all_notes:
            return

        # Store in SharedDataStore
        if HCR20_EXTRACTOR_AVAILABLE:
            try:
                shared_store = get_shared_store()
                shared_store.set_notes(all_notes, source="hcr20_import")
            except Exception as e:
                print(f"[HCR-20] Failed to store in SharedDataStore: {e}")

        # Store locally as well
        self._extracted_raw_notes = all_notes

        # Update extractor
        if self._hcr20_extractor:
            self._hcr20_extractor.set_notes(all_notes)

        # Show success message
        from PySide6.QtWidgets import QMessageBox
        msg = f"Successfully imported {len(all_notes)} note entries.\n\n"
        msg += "Extracting relevant data for all HCR-20 items...\n\n"
        msg += "Extraction searches:\n"
        msg += "- H items: ALL historical notes\n"
        msg += "- C items: Last 6 months only\n"
        msg += "- R items: Future-oriented content"

        QMessageBox.information(self, "Import Complete", msg)

        # Automatically extract and populate HCR-20 items
        self._extract_hcr20_from_notes()

    def _ingest_file(self, file_path: str) -> list:
        """Ingest a single file and return list of note entries.

        Uses the proper importer system that correctly parses dates from notes.
        """
        import os

        ext = os.path.splitext(file_path)[1].lower()
        fname = os.path.basename(file_path)

        # For Excel files, use the proper autodetect importer which parses dates correctly
        if ext in ['.xls', '.xlsx']:
            return self._ingest_excel_with_dates(file_path)
        elif ext == '.pdf':
            return self._ingest_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._ingest_docx(file_path)
        elif ext == '.txt':
            return self._ingest_txt(file_path)
        else:
            # Try as text
            return self._ingest_txt(file_path)

    def _ingest_pdf(self, file_path: str) -> list:
        """Ingest a PDF file."""
        try:
            import fitz  # PyMuPDF
            import os
            from datetime import datetime

            fname = os.path.basename(file_path)
            doc = fitz.open(file_path)
            notes = []

            for i, page in enumerate(doc):
                text = page.get_text().strip()
                if text:
                    notes.append({
                        'content': text,
                        'source': fname,
                        'page': i + 1,
                        'date': datetime.now(),
                        'type': 'pdf'
                    })

            doc.close()
            return notes
        except ImportError:
            # PyMuPDF not available, try alternative
            return self._ingest_pdf_fallback(file_path)
        except Exception as e:
            print(f"[HCR-20] PDF ingestion error: {e}")
            return []

    def _ingest_pdf_fallback(self, file_path: str) -> list:
        """Fallback PDF ingestion using pdfplumber or other methods."""
        try:
            import pdfplumber
            import os
            from datetime import datetime

            fname = os.path.basename(file_path)
            notes = []

            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        notes.append({
                            'content': text.strip(),
                            'source': fname,
                            'page': i + 1,
                            'date': datetime.now(),
                            'type': 'pdf'
                        })
            return notes
        except ImportError:
            print("[HCR-20] Neither PyMuPDF nor pdfplumber available for PDF ingestion")
            return []
        except Exception as e:
            print(f"[HCR-20] PDF fallback ingestion error: {e}")
            return []

    def _ingest_docx(self, file_path: str) -> list:
        """Ingest a Word document."""
        try:
            from docx import Document
            import os
            from datetime import datetime

            fname = os.path.basename(file_path)
            doc = Document(file_path)
            notes = []

            # Extract all paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Extract from tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())

            if full_text:
                notes.append({
                    'content': '\n\n'.join(full_text),
                    'source': fname,
                    'date': datetime.now(),
                    'type': 'docx'
                })

            return notes
        except Exception as e:
            print(f"[HCR-20] DOCX ingestion error: {e}")
            return []

    def _ingest_txt(self, file_path: str) -> list:
        """Ingest a text file."""
        try:
            import os
            from datetime import datetime

            fname = os.path.basename(file_path)

            # Try different encodings
            content = None
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue

            if content and content.strip():
                return [{
                    'content': content.strip(),
                    'source': fname,
                    'date': datetime.now(),
                    'type': 'txt'
                }]
            return []
        except Exception as e:
            print(f"[HCR-20] TXT ingestion error: {e}")
            return []

    def _ingest_excel_with_dates(self, file_path: str) -> list:
        """Ingest an Excel file using the proper autodetect importer that parses dates correctly.

        Uses importer_autodetect which handles RiO, CareNotes, and EPJS formats
        and properly extracts dates from clinical notes.
        """
        try:
            from importer_autodetect import import_files_autodetect
            import os

            fname = os.path.basename(file_path)
            print(f"[HCR-20] Using autodetect importer for {fname}")

            # Use the proper importer that parses dates from notes
            notes = import_files_autodetect([file_path])

            # Ensure notes have required fields for HCR-20 extraction
            for note in notes:
                # Normalize field names
                if 'text' not in note and 'content' in note:
                    note['text'] = note['content']
                if 'content' not in note and 'text' in note:
                    note['content'] = note['text']
                # Add source info
                note['source'] = note.get('source', fname)
                note['source_file'] = note.get('source_file', file_path)

            print(f"[HCR-20] Imported {len(notes)} dated entries from {fname}")
            return notes

        except ImportError as e:
            print(f"[HCR-20] Autodetect importer not available: {e}")
            # Fallback to basic Excel import
            return self._ingest_excel_basic(file_path)
        except Exception as e:
            print(f"[HCR-20] Excel ingestion error: {e}")
            import traceback
            traceback.print_exc()
            return []

    def _ingest_excel_basic(self, file_path: str) -> list:
        """Basic Excel ingestion fallback (without proper date parsing)."""
        try:
            import pandas as pd
            import os
            from datetime import datetime

            fname = os.path.basename(file_path)
            notes = []

            # Read all sheets
            xl = pd.ExcelFile(file_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Convert to text
                text_parts = []
                for _, row in df.iterrows():
                    row_text = ' | '.join([str(v) for v in row.values if pd.notna(v) and str(v).strip()])
                    if row_text:
                        text_parts.append(row_text)

                if text_parts:
                    notes.append({
                        'content': '\n'.join(text_parts),
                        'text': '\n'.join(text_parts),
                        'source': f"{fname} - {sheet_name}",
                        'date': datetime.now(),
                        'type': 'excel'
                    })

            return notes
        except Exception as e:
            print(f"[HCR-20] Basic Excel ingestion error: {e}")
            return []

    # ================================================================
    # HCR-20 EXTRACTION FROM NOTES
    # ================================================================

    def _connect_shared_store(self):
        """Connect to SharedDataStore for notes updates."""
        if not HCR20_EXTRACTOR_AVAILABLE:
            print("[HCR-20] Extractor not available - SharedDataStore not connected")
            return

        try:
            shared_store = get_shared_store()
            shared_store.notes_changed.connect(self._on_notes_changed)
            print("[HCR-20] Connected to SharedDataStore")
        except Exception as e:
            print(f"[HCR-20] Failed to connect to SharedDataStore: {e}")

    def _on_notes_changed(self, notes: list):
        """Handle when notes are updated in SharedDataStore."""
        if not notes:
            return

        self._extracted_raw_notes = notes
        print(f"[HCR-20] Received {len(notes)} notes from SharedDataStore")

        # Update the extractor with new notes
        if self._hcr20_extractor:
            self._hcr20_extractor.set_notes(notes)

    def _extract_hcr20_from_notes(self):
        """Extract HCR-20 relevant data from uploaded notes."""
        if not HCR20_EXTRACTOR_AVAILABLE:
            QMessageBox.warning(
                self, "Extractor Unavailable",
                "The HCR-20 extractor module is not available."
            )
            return

        # Get notes from SharedDataStore
        try:
            shared_store = get_shared_store()
            notes = shared_store.notes
        except:
            notes = self._extracted_raw_notes

        if not notes:
            QMessageBox.warning(
                self, "No Notes Available",
                "Please upload clinical notes first using the Notes panel or Import function.\n\n"
                "The extractor will search:\n"
                "- H items (Historical): ALL notes\n"
                "- C items (Clinical): Last 6 months only\n"
                "- R items (Risk Management): Future-oriented content"
            )
            return

        # Set up the extractor
        if not self._hcr20_extractor:
            self._hcr20_extractor = HCR20Extractor()

        self._hcr20_extractor.set_notes(notes)
        self._extracted_raw_notes = notes  # Ensure raw notes available for imported data popups

        # Run extraction for all items with progress feedback
        from PySide6.QtWidgets import QProgressDialog
        from PySide6.QtCore import Qt

        # Create progress dialog
        progress = QProgressDialog("Extracting data from notes...", None, 0, 100, self)
        progress.setWindowTitle("HCR-20 Extraction")
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.setMinimumDuration(0)
        progress.setValue(0)
        progress.show()
        QApplication.processEvents()

        try:
            # Step 1: Run main extraction (40%)
            progress.setLabelText("Step 1/3: Searching notes for HCR-20 terms...")
            progress.setValue(5)
            QApplication.processEvents()

            results = self._hcr20_extractor.extract_all()
            summary = self._hcr20_extractor.get_summary()
            progress.setValue(40)
            QApplication.processEvents()

            # Step 2: Populate form fields (20%)
            progress.setLabelText("Step 2/3: Populating form fields...")
            QApplication.processEvents()

            items_populated = 0
            for item_key in results.keys():
                if self._populate_hcr_item_from_extraction(item_key.lower(), results[item_key]):
                    items_populated += 1
            progress.setValue(60)
            QApplication.processEvents()

            # Step 3: Run risk analysis ONCE and cache it (40%)
            progress.setLabelText("Step 3/3: Analyzing risk patterns...")
            QApplication.processEvents()

            # Cache risk analysis results to avoid running 3 times
            from risk_overview_panel import analyze_notes_for_risk
            # Limit notes for risk analysis to improve performance
            risk_notes = notes[:2000] if len(notes) > 2000 else notes
            cached_risk_results = analyze_notes_for_risk(risk_notes)
            progress.setValue(80)
            QApplication.processEvents()

            # H1 and H2 forensic data already populated by _populate_imported_entries_forensic()
            # (called from _populate_hcr_item_from_extraction in Step 2 above)

            progress.setValue(90)
            QApplication.processEvents()

            # H5 substance data already populated by _populate_imported_entries_all_notes()
            # (called from _populate_hcr_item_from_extraction in Step 2 above)

            progress.setValue(100)
            progress.close()

            # Show summary
            msg = (
                f"Extraction Complete\n\n"
                f"Notes searched: {summary['total_notes']}\n"
                f"Items with matches: {summary['items_with_matches']}\n"
                f"Items without matches: {summary['items_without_matches']}\n\n"
                f"Historical (H1-H10): {summary['by_scope']['historical']['with_matches']} items with data\n"
                f"Clinical (C1-C5, last 6 months): {summary['by_scope']['clinical']['with_matches']} items with data\n"
                f"Risk Management (R1-R5): {summary['by_scope']['risk_management']['with_matches']} items with data\n\n"
                f"Review each section and edit as needed."
            )
            QMessageBox.information(self, "Extraction Complete", msg)

        except Exception as e:
            QMessageBox.critical(
                self, "Extraction Error",
                f"An error occurred during extraction:\n{str(e)}"
            )
            import traceback
            traceback.print_exc()

    def _populate_hcr_item_from_extraction(self, key: str, extraction_result: dict) -> bool:
        """
        Populate an HCR-20 item's popup fields from extraction results.

        Args:
            key: The item key (e.g., 'h1', 'c3', 'r5')
            extraction_result: The extraction result from HCR20Extractor

        Returns:
            True if any data was populated
        """
        if not extraction_result or 'error' in extraction_result:
            return False

        populated = False

        # Populate the Imported Data section with collapsible dated entries
        # For H1 and H2, use forensic-style categorization (like Section 5 in tribunal reports)
        if key in ('h1', 'h2'):
            entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
            if entries_layout:
                populated = self._populate_imported_entries_forensic(key, entries_layout)
                # Expand the collapsible section if data was populated
                if populated:
                    imported_section = getattr(self, f"popup_{key}_imported_section", None)
                    if imported_section and hasattr(imported_section, '_is_collapsed'):
                        if imported_section._is_collapsed:
                            imported_section._toggle_collapse()

        # For H3, use separate imported data sections for intimate and non-intimate relationships
        elif key == 'h3':
            populated = self._populate_h3_imported_entries(extraction_result)

        # For H4, use separate imported data sections for education and employment
        elif key == 'h4':
            populated = self._populate_h4_imported_entries(extraction_result)

        # For H5, use the EXACT same search as the Notes Risk Panel Substance Misuse section
        elif key == 'h5':
            entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
            if entries_layout:
                populated = self._populate_h5_substance_risk_search(entries_layout)
                if populated:
                    imported_section = getattr(self, f"popup_{key}_imported_section", None)
                    if imported_section and hasattr(imported_section, '_is_collapsed'):
                        if imported_section._is_collapsed:
                            imported_section._toggle_collapse()

        # For H6-H10, search ALL raw notes directly (bypass 2000-note limit)
        elif key in ('h6', 'h7', 'h8', 'h9', 'h10'):
            entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
            if entries_layout:
                populated = self._populate_imported_entries_all_notes(key, entries_layout)
                if populated:
                    imported_section = getattr(self, f"popup_{key}_imported_section", None)
                    if imported_section and hasattr(imported_section, '_is_collapsed'):
                        if imported_section._is_collapsed:
                            imported_section._toggle_collapse()

        # For C1-C5, search ALL raw notes for last 6 months of clinical data
        elif key in ('c1', 'c2', 'c3', 'c4', 'c5'):
            entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
            if entries_layout:
                populated = self._populate_clinical_from_all_notes(key, entries_layout)
                if populated:
                    imported_section = getattr(self, f"popup_{key}_imported_section", None)
                    if imported_section and hasattr(imported_section, '_is_collapsed'):
                        if imported_section._is_collapsed:
                            imported_section._toggle_collapse()

        # For R1-R5, use standard import from extractor results
        else:
            entries_layout = getattr(self, f"popup_{key}_imported_entries_layout", None)
            if entries_layout:
                populated = self._populate_imported_entries(key, extraction_result, entries_layout)
                if populated:
                    imported_section = getattr(self, f"popup_{key}_imported_section", None)
                    if imported_section and hasattr(imported_section, '_is_collapsed'):
                        if imported_section._is_collapsed:
                            imported_section._toggle_collapse()

        # Get suggested ratings
        suggested = self._hcr20_extractor.get_suggested_rating(key.upper())

        # Try to set presence rating
        presence_suggestion = suggested.get('presence', '')
        if presence_suggestion:
            presence_group = getattr(self, f"popup_{key}_presence_group", None)
            if presence_group:
                if 'Absent' in presence_suggestion:
                    btn = getattr(self, f"popup_{key}_presence_no", None)
                elif 'Partial' in presence_suggestion or 'Possibly' in presence_suggestion:
                    btn = getattr(self, f"popup_{key}_presence_partial", None)
                elif 'Present' in presence_suggestion:
                    btn = getattr(self, f"popup_{key}_presence_yes", None)
                else:
                    btn = None

                if btn:
                    btn.setChecked(True)
                    populated = True

        # Try to set relevance rating
        relevance_suggestion = suggested.get('relevance', '')
        if relevance_suggestion:
            relevance_group = getattr(self, f"popup_{key}_relevance_group", None)
            if relevance_group:
                if 'Low' in relevance_suggestion:
                    btn = getattr(self, f"popup_{key}_relevance_low", None)
                elif 'Moderate' in relevance_suggestion:
                    btn = getattr(self, f"popup_{key}_relevance_mod", None)
                elif 'High' in relevance_suggestion:
                    btn = getattr(self, f"popup_{key}_relevance_high", None)
                else:
                    btn = None

                if btn:
                    btn.setChecked(True)
                    populated = True

        # Populate subsection fields if they exist
        # SKIP for H3 and H4 - they use separate imported data sections and user ticks to fill subsections
        if key not in ('h3', 'h4'):
            subsection_matches = extraction_result.get('subsection_matches', {})
            for subsection_key, matches in subsection_matches.items():
                # Try to find a matching text field for this subsection
                safe_name = subsection_key.lower().replace(" ", "_").replace("/", "_").replace(":", "").replace("-", "_")
                subsection_field = getattr(self, f"popup_{key}_{safe_name}", None)

                if subsection_field and matches:
                    # Format the matches for this subsection - no limit for historical items
                    lines = []
                    for match in matches:
                        date = match.get('date')
                        date_str = date.strftime('%d/%m/%Y') if date else 'Unknown date'
                        for term_match in match.get('matches', [])[:2]:
                            lines.append(f"[{date_str}] {term_match['excerpt'][:300]}...")

                    if lines:
                        subsection_field.setPlainText("\n\n".join(lines))
                        populated = True

        return populated

    def _populate_imported_entries(self, key: str, extraction_result: dict, entries_layout) -> bool:
        """
        Populate the imported data section with collapsible dated entry boxes.

        Args:
            key: The item key (e.g., 'h1', 'c3')
            extraction_result: The extraction result from HCR20Extractor
            entries_layout: The QVBoxLayout to add entries to

        Returns:
            True if any entries were added
        """
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy
        from PySide6.QtCore import Qt

        # Clear existing entries
        checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            item = entries_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Get matches from extraction result
        main_matches = extraction_result.get('main_matches', [])
        if not main_matches:
            return False

        # Sort by date (newest first)
        def get_sort_date(match):
            dt = match.get('date')
            if dt is None:
                return ""
            if hasattr(dt, "strftime"):
                return dt.strftime("%Y-%m-%d")
            return str(dt)

        sorted_matches = sorted(main_matches, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep only the most relevant entry per date
        # (the one with the most term matches or longest excerpt)
        seen_dates = {}
        deduplicated_matches = []
        for match in sorted_matches:
            date_key = get_sort_date(match)

            # Calculate relevance score: number of matches + total excerpt length
            num_matches = len(match.get('matches', []))
            total_excerpt_len = sum(len(tm.get('excerpt', '')) for tm in match.get('matches', []))
            relevance_score = num_matches * 100 + total_excerpt_len

            if date_key not in seen_dates:
                seen_dates[date_key] = (match, relevance_score)
                deduplicated_matches.append(match)
            else:
                # Check if this match is more relevant than the stored one
                existing_match, existing_score = seen_dates[date_key]
                if relevance_score > existing_score:
                    # Replace the less relevant entry
                    deduplicated_matches.remove(existing_match)
                    deduplicated_matches.append(match)
                    seen_dates[date_key] = (match, relevance_score)

        # Re-sort after deduplication
        deduplicated_matches = sorted(deduplicated_matches, key=get_sort_date, reverse=True)

        # For clinical items (C1-C5), filter to only the last 6 months of notes
        # This reflects current clinical status rather than historical patterns
        if key in ('c1', 'c2', 'c3', 'c4', 'c5') and deduplicated_matches:
            from datetime import datetime, timedelta
            from dateutil.relativedelta import relativedelta

            # Find the most recent date in the notes
            most_recent_date = None
            for match in deduplicated_matches:
                dt = match.get('date')
                if dt:
                    if hasattr(dt, 'date'):
                        dt = dt.date() if hasattr(dt, 'date') else dt
                    if isinstance(dt, str):
                        try:
                            dt = datetime.strptime(dt, "%Y-%m-%d").date()
                        except:
                            continue
                    if most_recent_date is None or dt > most_recent_date:
                        most_recent_date = dt

            if most_recent_date:
                # Calculate cutoff date (6 months before most recent)
                if hasattr(most_recent_date, 'date'):
                    most_recent_date = most_recent_date
                cutoff_date = most_recent_date - relativedelta(months=6)

                # Filter matches to only those within the 6-month window
                filtered_matches = []
                for match in deduplicated_matches:
                    dt = match.get('date')
                    if dt:
                        if hasattr(dt, 'date'):
                            match_date = dt.date() if hasattr(dt.date, '__call__') else dt
                        elif isinstance(dt, str):
                            try:
                                match_date = datetime.strptime(dt, "%Y-%m-%d").date()
                            except:
                                filtered_matches.append(match)  # Keep if can't parse
                                continue
                        else:
                            match_date = dt

                        if match_date >= cutoff_date:
                            filtered_matches.append(match)
                    else:
                        # Keep entries without dates (shouldn't happen often)
                        filtered_matches.append(match)

                deduplicated_matches = filtered_matches
                print(f"[HCR20] {key.upper()}: Filtered to {len(deduplicated_matches)} entries within 6 months of {most_recent_date}")

        # For historical items (H3-H10), show ALL entries - no limit
        # For clinical (C1-C5) and risk (R1-R5), limit to 30 entries
        scope = extraction_result.get('scope', 'historical')
        max_entries = 999 if scope == 'historical' else 30

        entries_added = 0
        for match in deduplicated_matches[:max_entries]:
            dt = match.get('date')
            note = match.get('note', {})
            source = note.get('source', '') if isinstance(note, dict) else ''

            # For H6-H10, C1-C5, R1-R5, show excerpts around the matched terms
            if key in ('h6', 'h7', 'h8', 'h10', 'c1', 'c2', 'c3', 'c4', 'c5', 'r1', 'r2', 'r3', 'r4', 'r5'):
                match_texts = []
                for term_match in match.get('matches', []):
                    excerpt = term_match.get('excerpt', '').strip()
                    if excerpt:
                        match_texts.append(excerpt)
                text = "\n\n".join(match_texts) if match_texts else ''
            else:
                # Get the FULL note content (not just excerpts)
                if isinstance(note, dict):
                    text = note.get('text', '') or note.get('content', '') or ''
                else:
                    text = str(note) if note else ''

                # If no full text available, fall back to matched excerpts
                if not text.strip():
                    match_texts = []
                    for term_match in match.get('matches', []):
                        excerpt = term_match.get('excerpt', '').strip()
                        if excerpt:
                            match_texts.append(excerpt)
                    text = "\n\n".join(match_texts)

            if not text.strip():
                continue

            # Collect matched terms for highlighting
            matched_terms = set()
            for term_match in match.get('matches', []):
                term = term_match.get('term', '')
                if term:
                    matched_terms.add(term.lower())

            # Create highlighted HTML version of text
            import html
            import re
            escaped_text = html.escape(text)

            # Highlight matched terms in the text (yellow background)
            highlighted_text = escaped_text
            for term in matched_terms:
                # Create case-insensitive pattern for the term
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                highlighted_text = pattern.sub(
                    lambda m: f'<span style="background-color: #ffeb3b; padding: 1px 3px; border-radius: 3px; font-weight: 600;">{m.group()}</span>',
                    highlighted_text
                )

            # Convert newlines to HTML breaks
            highlighted_html = highlighted_text.replace('\n', '<br>')

            # Format date
            if dt:
                if hasattr(dt, "strftime"):
                    date_str = dt.strftime("%d %b %Y")
                else:
                    date_str = str(dt)
            else:
                date_str = "No date"

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("entryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet("""
                QFrame#entryFrame {
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-radius: 8px;
                    padding: 4px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row with toggle, date, source, and checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(180, 150, 50, 0.2);
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: #806000;
                }
                QPushButton:hover { background: rgba(180, 150, 50, 0.35); }
            """)
            header_row.addWidget(toggle_btn)

            date_label = QLabel(f"{date_str}")
            date_label.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            if source:
                source_label = QLabel(f"({source})")
                source_label.setStyleSheet("""
                    QLabel {
                        font-size: 14px;
                        color: #666;
                        background: transparent;
                        border: none;
                    }
                """)
                header_row.addWidget(source_label)

            header_row.addStretch()

            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body text (hidden by default) - with highlighted matched terms
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet("""
                QTextEdit {
                    font-size: 15px;
                    color: #333;
                    background: rgba(255, 248, 220, 0.5);
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }
            """)
            body_text.setVisible(False)

            # Calculate height based on content
            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))  # Reduced by 10% (was 160)

            # Toggle function
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("▸")
                    else:
                        body.setVisible(True)
                        btn.setText("▾")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout.addWidget(body_text)
            entries_layout.addWidget(entry_frame)
            entries_added += 1

        # Store checkboxes reference
        setattr(self, f"popup_{key}_imported_checkboxes", checkboxes)

        return entries_added > 0

    def _populate_h5_substance_risk_search(self, entries_layout) -> bool:
        """Populate H5 imported data using the EXACT same search process as the
        Notes Risk Panel Substance Misuse section in risk_overview_panel.py.
        Uses RISK_CATEGORIES patterns, is_false_positive(), and _has_negative_context()."""
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy, QWidget
        from PySide6.QtCore import Qt
        import html
        import re

        all_notes = getattr(self, '_extracted_raw_notes', [])
        if not all_notes:
            print("[HCR-20 H5] No raw notes available for substance risk search")
            return False

        # Import the exact risk panel search functions and patterns
        try:
            from risk_overview_panel import RISK_CATEGORIES, is_false_positive, _has_negative_context, _normalise_date
        except ImportError as e:
            print(f"[HCR-20 H5] Could not import risk_overview_panel functions: {e}")
            return False

        substance_config = RISK_CATEGORIES.get("Substance Misuse")
        if not substance_config or "subcategories" not in substance_config:
            print("[HCR-20 H5] No Substance Misuse config found in RISK_CATEGORIES")
            return False

        subcategories = substance_config["subcategories"]
        accent_color = substance_config.get("color", "#9c27b0")

        print(f"[HCR-20 H5] Searching ALL {len(all_notes)} raw notes using Risk Panel Substance Misuse patterns")

        # Run the exact same search logic as analyze_notes_for_risk() but only for Substance Misuse
        all_matches = []
        for note in all_notes:
            if not isinstance(note, dict):
                continue

            text = note.get("text", "") or note.get("content", "") or note.get("body", "")
            if not text or not text.strip():
                continue

            text_lower = text.lower()
            note_date = _normalise_date(note.get("date") or note.get("datetime"))
            matched_in_note = set()

            for subcat_name, subcat_config in subcategories.items():
                for pattern in subcat_config["patterns"]:
                    try:
                        match = re.search(pattern, text_lower)
                        if match:
                            match_key = ("Substance Misuse", match.group())
                            if match_key in matched_in_note:
                                continue

                            if is_false_positive(text, match.start()):
                                continue

                            if _has_negative_context(text_lower, match.start(), match.end()):
                                continue

                            matched_in_note.add(match_key)
                            severity = subcat_config["severity"]

                            all_matches.append({
                                'text': text,
                                'date': note_date,
                                'subcategory': subcat_name,
                                'severity': severity,
                                'matched': match.group(),
                                'subcat_color': subcat_config.get("color", accent_color),
                                'patterns': subcat_config["patterns"],
                            })
                            break  # Move to next subcategory after finding a match
                    except Exception:
                        pass

        if not all_matches:
            print("[HCR-20 H5] No substance misuse matches found")
            return False

        # Sort by date (newest first)
        def get_sort_date(m):
            dt = m.get('date')
            if dt is None:
                return ""
            if hasattr(dt, "strftime"):
                return dt.strftime("%Y-%m-%d")
            return str(dt)

        sorted_matches = sorted(all_matches, key=get_sort_date, reverse=True)

        print(f"[HCR-20 H5] Found {len(sorted_matches)} substance misuse incidents using risk panel search")

        # Clear existing entries
        key = 'h5'
        checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Create incidents container
        incidents_container = QWidget()
        incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(8)
        entries_layout.addWidget(incidents_container)

        bg_color = f"rgba({int(accent_color[1:3], 16)}, {int(accent_color[3:5], 16)}, {int(accent_color[5:7], 16)}, 0.08)"
        border_color = f"{accent_color}66"

        for match_data in sorted_matches:
            date = match_data['date']
            text = match_data['text']
            subcat = match_data['subcategory']
            severity = match_data['severity']
            matched_text = match_data['matched']
            subcat_color = match_data['subcat_color']
            patterns = match_data['patterns']

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Create highlighted HTML using the risk panel's highlight approach
            escaped_text = html.escape(text)
            highlighted_html = escaped_text
            for pat in patterns:
                try:
                    for m in re.finditer(pat, text, re.IGNORECASE):
                        escaped_match = html.escape(m.group())
                        escaped_pattern = re.compile(re.escape(escaped_match), re.IGNORECASE)
                        highlighted_html = escaped_pattern.sub(
                            f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_match}</span>',
                            highlighted_html
                        )
                except Exception:
                    pass
            highlighted_html = highlighted_html.replace('\n', '<br>')

            # Severity color
            severity_colors = {'high': '#dc2626', 'medium': '#d97706', 'low': '#059669'}
            sev_color = severity_colors.get(severity, '#6b7280')

            # Entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("h5SubstanceEntry")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#h5SubstanceEntry {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {subcat_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_layout_v = QVBoxLayout(entry_frame)
            entry_layout_v.setContentsMargins(10, 8, 10, 8)
            entry_layout_v.setSpacing(6)
            entry_layout_v.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button
            toggle_btn = QPushButton("\u25b8")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(date_str)
            date_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 16px;
                    font-weight: 600;
                    color: {accent_color};
                    background: transparent;
                    border: none;
                }}
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Subcategory badge
            subcat_badge = QLabel(subcat)
            subcat_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 11px;
                    font-weight: 600;
                    color: white;
                    background: {subcat_color}cc;
                    border: none;
                    border-radius: 3px;
                    padding: 2px 6px;
                }}
            """)
            header_row.addWidget(subcat_badge)

            # Severity badge
            sev_badge = QLabel(severity.upper())
            sev_badge.setStyleSheet(f"""
                QLabel {{
                    font-size: 10px;
                    font-weight: 700;
                    color: white;
                    background: {sev_color};
                    border: none;
                    border-radius: 3px;
                    padding: 2px 5px;
                }}
            """)
            header_row.addWidget(sev_badge)

            header_row.addStretch()

            # Checkbox
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(lambda _, k='h5': self._update_preview(k))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout_v.addLayout(header_row)

            # Body text (hidden by default)
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"""
                QTextEdit {{
                    font-size: 15px;
                    color: #333;
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }}
            """)
            body_text.setVisible(False)

            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))

            # Toggle function
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("\u25b8")
                    else:
                        body.setVisible(True)
                        btn.setText("\u25be")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout_v.addWidget(body_text)
            incidents_layout.addWidget(entry_frame)

        # Store checkboxes reference
        setattr(self, f"popup_{key}_imported_checkboxes", checkboxes)

        print(f"[HCR-20 H5] Displayed {len(sorted_matches)} substance misuse entries using risk panel search")
        return len(sorted_matches) > 0

    def _populate_clinical_from_all_notes(self, key, entries_layout) -> bool:
        """Populate C1-C5 imported data by searching ALL raw notes for relevant clinical data,
        filtered to the most recent 6 months working back from the most recent entry."""
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy, QWidget
        from PySide6.QtCore import Qt
        from datetime import datetime, timedelta
        import html
        import re

        key_upper = key.upper()
        all_notes = getattr(self, '_extracted_raw_notes', [])
        if not all_notes:
            print(f"[HCR-20 {key_upper}] No raw notes available for clinical search")
            return False

        # Get terms from extractor config
        try:
            from hcr20_extractor import HCR20_EXTRACTION_TERMS
            item_config = HCR20_EXTRACTION_TERMS.get(key_upper, {})
            terms = item_config.get('terms', [])
            subsections_config = item_config.get('subsections', {})
        except Exception as e:
            print(f"[HCR-20 {key_upper}] Could not load extraction terms: {e}")
            return False

        if not terms:
            print(f"[HCR-20 {key_upper}] No terms defined")
            return False

        # Build term-to-subsection mapping (auto-generate labels from keys)
        term_to_subsection = {}
        for sub_key, sub_terms in subsections_config.items():
            label = sub_key.replace('_', ' ').title()
            for t in sub_terms:
                term_to_subsection[t.lower()] = label

        # Search all notes for matching terms
        all_matches = []
        for note in all_notes:
            if not isinstance(note, dict):
                continue

            text = note.get('text', '') or note.get('body', '') or note.get('content', '')
            if not text or not text.strip():
                continue

            text_lower = text.lower()
            matched_terms = []
            matched_categories = set()
            for term in terms:
                if term.lower() in text_lower:
                    matched_terms.append(term)
                    cat = term_to_subsection.get(term.lower(), '')
                    if cat:
                        matched_categories.add(cat)

            if not matched_terms:
                continue

            note_date = note.get('date') or note.get('datetime')
            # Normalize date
            dt = None
            if note_date:
                if isinstance(note_date, datetime):
                    dt = note_date
                elif hasattr(note_date, 'to_pydatetime'):
                    try:
                        dt = note_date.to_pydatetime()
                    except Exception:
                        pass
                else:
                    try:
                        import pandas as pd
                        dt = pd.to_datetime(note_date, errors='coerce', dayfirst=True)
                        if pd.isna(dt):
                            dt = None
                        else:
                            dt = dt.to_pydatetime()
                    except Exception:
                        pass

            all_matches.append({
                'text': text,
                'date': dt,
                'matched_terms': matched_terms,
                'categories': list(matched_categories),
            })

        if not all_matches:
            print(f"[HCR-20 {key_upper}] No matches found in all notes")
            return False

        # Find most recent date and filter to last 6 months
        dates_with_values = [m['date'] for m in all_matches if m['date'] is not None]
        if dates_with_values:
            most_recent = max(dates_with_values)
            cutoff = most_recent - timedelta(days=183)  # ~6 months
            filtered = [m for m in all_matches if m['date'] is None or m['date'] >= cutoff]
            print(f"[HCR-20 {key_upper}] Most recent note: {most_recent.strftime('%d %b %Y')}, cutoff: {cutoff.strftime('%d %b %Y')}")
            print(f"[HCR-20 {key_upper}] {len(all_matches)} total matches -> {len(filtered)} within last 6 months")
        else:
            filtered = all_matches
            print(f"[HCR-20 {key_upper}] No dated entries, showing all {len(filtered)} matches")

        # Sort by date (newest first)
        def get_sort_date(m):
            dt = m.get('date')
            if dt is None:
                return ""
            return dt.strftime("%Y-%m-%d")

        sorted_matches = sorted(filtered, key=get_sort_date, reverse=True)

        # Deduplicate by date
        seen_dates = {}
        deduplicated = []
        for match in sorted_matches:
            date_key = get_sort_date(match)
            num_terms = len(match.get('matched_terms', []))
            text_len = len(match.get('text', ''))
            relevance = num_terms * 100 + text_len

            if date_key not in seen_dates or relevance > seen_dates[date_key][1]:
                if date_key in seen_dates:
                    old_match = seen_dates[date_key][0]
                    if old_match in deduplicated:
                        deduplicated.remove(old_match)
                seen_dates[date_key] = (match, relevance)
                deduplicated.append(match)

        sorted_matches = sorted(deduplicated, key=get_sort_date, reverse=True)
        print(f"[HCR-20 {key_upper}] Displaying {len(sorted_matches)} entries (deduplicated)")

        # Clear existing entries
        checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Create incidents container
        incidents_container = QWidget()
        incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(8)
        entries_layout.addWidget(incidents_container)

        CLINICAL_COLORS = {
            'c1': '#2563eb',  # Blue
            'c2': '#dc2626',  # Red
            'c3': '#7c3aed',  # Violet
            'c4': '#d97706',  # Amber
            'c5': '#059669',  # Green
        }
        accent_color = CLINICAL_COLORS.get(key, '#2563eb')
        bg_color = f"rgba({int(accent_color[1:3], 16)}, {int(accent_color[3:5], 16)}, {int(accent_color[5:7], 16)}, 0.08)"
        border_color = f"{accent_color}66"

        # Filter bar
        filterable_frames = []
        all_cats = set()
        for m in sorted_matches:
            for c in m.get('categories', []):
                all_cats.add(c)

        if all_cats:
            filter_frame = QFrame()
            filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
            filter_row = QHBoxLayout(filter_frame)
            filter_row.setContentsMargins(0, 0, 0, 4)
            filter_row.setSpacing(4)

            flbl = QLabel("Filter:")
            flbl.setStyleSheet("QLabel { font-size: 12px; font-weight: 600; color: #374151; background: transparent; border: none; }")
            filter_row.addWidget(flbl)

            def make_filter(cat_name, flist):
                def do_filter():
                    for frm, cats in flist:
                        frm.setVisible(cat_name in cats)
                return do_filter

            def make_show_all(flist):
                def do_show():
                    for frm, cats in flist:
                        frm.setVisible(True)
                return do_show

            sa_btn = QPushButton("Show All")
            sa_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            sa_btn.setStyleSheet("QPushButton { font-size: 11px; font-weight: 600; color: #1f2937; background-color: #e5e7eb; border: 1px solid #d1d5db; border-radius: 3px; padding: 3px 8px; } QPushButton:hover { background-color: #d1d5db; }")
            sa_btn.clicked.connect(make_show_all(filterable_frames))
            filter_row.addWidget(sa_btn)

            for cat in sorted(all_cats):
                cb_btn = QPushButton(cat)
                cb_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                cb_btn.setStyleSheet(f"QPushButton {{ font-size: 11px; font-weight: 600; color: {accent_color}; background-color: {bg_color}; border: 1px solid {accent_color}66; border-radius: 3px; padding: 3px 8px; }} QPushButton:hover {{ background-color: {accent_color}22; }}")
                cb_btn.clicked.connect(make_filter(cat, filterable_frames))
                filter_row.addWidget(cb_btn)

            filter_row.addStretch()
            incidents_layout.addWidget(filter_frame)

        for match in sorted_matches:
            date = match['date']
            text = match['text']
            matched_terms = match.get('matched_terms', [])
            categories = match.get('categories', [])

            if not text or not text.strip():
                continue

            # Format date
            if date:
                date_str = date.strftime("%d %b %Y")
            else:
                date_str = "No date"

            # Highlighted HTML
            escaped_text = html.escape(text)
            highlighted_html = escaped_text
            for term in matched_terms[:5]:
                try:
                    escaped_term = html.escape(term)
                    pattern = re.compile(re.escape(escaped_term), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_term}</span>',
                        highlighted_html
                    )
                except Exception:
                    pass
            highlighted_html = highlighted_html.replace('\n', '<br>')

            # Entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("clinicalEntry")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#clinicalEntry {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {accent_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_v = QVBoxLayout(entry_frame)
            entry_v.setContentsMargins(10, 8, 10, 8)
            entry_v.setSpacing(6)
            entry_v.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button
            toggle_btn = QPushButton("\u25b8")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color}; border: none; border-radius: 4px;
                    font-size: 17px; font-weight: bold; color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(date_str)
            date_label.setStyleSheet(f"QLabel {{ font-size: 16px; font-weight: 600; color: {accent_color}; background: transparent; border: none; }}")
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Category labels
            for cat_label in categories[:2]:
                cat_badge = QLabel(cat_label)
                cat_badge.setStyleSheet(f"QLabel {{ font-size: 11px; font-weight: 700; color: #1f2937; background-color: {accent_color}30; border: 1px solid {accent_color}66; border-radius: 3px; padding: 2px 6px; }}")
                header_row.addWidget(cat_badge)

            # Term badges
            for term in matched_terms[:3]:
                term_badge = QLabel(term)
                term_badge.setStyleSheet(f"QLabel {{ font-size: 11px; font-weight: 600; color: #1f2937; background-color: {accent_color}20; border: 1px solid {accent_color}55; border-radius: 3px; padding: 2px 6px; }}")
                header_row.addWidget(term_badge)

            header_row.addStretch()

            # Checkbox
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("QCheckBox { background: transparent; } QCheckBox::indicator { width: 16px; height: 16px; }")
            cb.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_v.addLayout(header_row)

            # Body text (hidden by default)
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"QTextEdit {{ font-size: 15px; color: #333; background: {bg_color}; border: none; border-radius: 4px; padding: 6px; }}")
            body_text.setVisible(False)

            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))

            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("\u25b8")
                    else:
                        body.setVisible(True)
                        btn.setText("\u25be")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_v.addWidget(body_text)
            incidents_layout.addWidget(entry_frame)
            filterable_frames.append((entry_frame, categories))

        setattr(self, f"popup_{key}_imported_checkboxes", checkboxes)
        print(f"[HCR-20 {key_upper}] Displayed {len(sorted_matches)} clinical entries from last 6 months")
        return len(sorted_matches) > 0

    def _populate_imported_entries_all_notes(self, key: str, entries_layout) -> bool:
        """Populate imported data for H6-H10 by searching ALL raw notes directly (H5 uses
        _populate_h5_substance_risk_search instead). Bypasses the 2000-note limit in the extractor. Uses term definitions from
        hcr20_extractor.HCR20_EXTRACTION_TERMS."""
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy
        from PySide6.QtCore import Qt
        import html
        import re

        all_notes = getattr(self, '_extracted_raw_notes', [])
        if not all_notes:
            print(f"[HCR-20 {key.upper()}] No raw notes available for all-notes search")
            return False

        # Get terms from extractor config
        try:
            from hcr20_extractor import HCR20_EXTRACTION_TERMS
            item_config = HCR20_EXTRACTION_TERMS.get(key.upper(), {})
            terms = item_config.get('terms', [])
        except Exception as e:
            print(f"[HCR-20 {key.upper()}] Could not load extraction terms: {e}")
            return False

        if not terms:
            print(f"[HCR-20 {key.upper()}] No terms defined")
            return False

        # Build term-to-subsection mapping for category labels
        subsections_config = item_config.get('subsections', {})
        term_to_subsection = {}
        SUBSECTION_LABELS = {
            'violence_in_home': 'Violence in Home', 'maltreatment': 'Maltreatment',
            'physical_trauma': 'Physical Trauma', 'ptsd_related': 'PTSD Related',
            'caregiver_disruption': 'Caregiver Disruption',
            'alcohol': 'Alcohol', 'drugs': 'Drugs', 'impact_on_risk': 'Impact on Risk',
            'intimate': 'Intimate', 'non_intimate': 'Non-intimate',
            'education': 'Education', 'employment': 'Employment',
            'violent_attitudes': 'Violent Attitudes', 'antisocial_attitudes': 'Antisocial',
        }
        for sub_key, sub_terms in subsections_config.items():
            label = SUBSECTION_LABELS.get(sub_key, sub_key.replace('_', ' ').title())
            for t in sub_terms:
                term_to_subsection[t.lower()] = label

        # Log date range
        dates = []
        for note in all_notes:
            if isinstance(note, dict):
                dt = note.get('date') or note.get('datetime')
                if dt:
                    dates.append(dt)
        if dates:
            print(f"[HCR-20 {key.upper()}] Searching ALL {len(all_notes)} raw notes (date range: {min(dates)} to {max(dates)})")
        else:
            print(f"[HCR-20 {key.upper()}] Searching ALL {len(all_notes)} raw notes")

        # Search all notes for matching terms
        all_matches = []
        for note in all_notes:
            if not isinstance(note, dict):
                continue

            text = note.get('text', '') or note.get('body', '') or note.get('content', '')
            if not text or not text.strip():
                continue

            text_lower = text.lower()
            matched_terms = []
            matched_categories = set()
            for term in terms:
                if term.strip().startswith('r\'') or term.strip().startswith('r"'):
                    # Regex pattern
                    try:
                        if re.search(term.strip("r'\""), text_lower):
                            matched_terms.append(term)
                            cat = term_to_subsection.get(term.lower(), '')
                            if cat:
                                matched_categories.add(cat)
                    except:
                        pass
                else:
                    if term.lower() in text_lower:
                        matched_terms.append(term)
                        cat = term_to_subsection.get(term.lower(), '')
                        if cat:
                            matched_categories.add(cat)

            if not matched_terms:
                continue

            note_date = note.get('date') or note.get('datetime')
            all_matches.append({
                'text': text,
                'categories': list(matched_categories),
                'date': note_date,
                'matched_terms': matched_terms,
            })

        if not all_matches:
            print(f"[HCR-20 {key.upper()}] No matches found in all notes")
            return False

        # Sort by date (newest first)
        def get_sort_date(match):
            dt = match.get('date')
            if dt is None:
                return ""
            if hasattr(dt, "strftime"):
                return dt.strftime("%Y-%m-%d")
            return str(dt)

        sorted_matches = sorted(all_matches, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep most relevant per date
        seen_dates = {}
        deduplicated = []
        for match in sorted_matches:
            date_key = get_sort_date(match)
            num_terms = len(match.get('matched_terms', []))
            text_len = len(match.get('text', ''))
            relevance = num_terms * 100 + text_len

            if date_key not in seen_dates or relevance > seen_dates[date_key][1]:
                if date_key in seen_dates:
                    old_match = seen_dates[date_key][0]
                    if old_match in deduplicated:
                        deduplicated.remove(old_match)
                seen_dates[date_key] = (match, relevance)
                deduplicated.append(match)

        sorted_matches = sorted(deduplicated, key=get_sort_date, reverse=True)

        print(f"[HCR-20 {key.upper()}] Found {len(sorted_matches)} matching notes (deduplicated by date)")

        # Clear existing entries
        checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            child = entries_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Create incidents container
        incidents_container = QWidget()
        incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(8)
        entries_layout.addWidget(incidents_container)

        # Use item-specific accent color
        ITEM_COLORS = {
            'h5': '#9c27b0',  # Purple
            'h6': '#2563eb',  # Blue
            'h7': '#dc2626',  # Red
            'h8': '#059669',  # Green
            'h9': '#d97706',  # Amber
            'h10': '#7c3aed',  # Violet
        }
        accent_color = ITEM_COLORS.get(key, '#806000')
        bg_color = f"rgba({int(accent_color[1:3], 16)}, {int(accent_color[3:5], 16)}, {int(accent_color[5:7], 16)}, 0.08)"
        border_color = f"{accent_color}66"

        # Filter bar for items with subcategories (H8, H10): buttons per category + Show All
        filterable_entry_frames = []  # list of (entry_frame, categories)
        if key in ('h8', 'h10'):
            # Collect all unique categories from matches
            all_cats = set()
            for m in sorted_matches:
                for c in m.get('categories', []):
                    all_cats.add(c)

            if all_cats:
                filter_frame = QFrame()
                filter_frame.setStyleSheet("QFrame { background: transparent; border: none; }")
                filter_layout = QHBoxLayout(filter_frame)
                filter_layout.setContentsMargins(0, 0, 0, 4)
                filter_layout.setSpacing(4)

                filter_lbl = QLabel("Filter:")
                filter_lbl.setStyleSheet("QLabel { font-size: 12px; font-weight: 600; color: #374151; background: transparent; border: none; }")
                filter_layout.addWidget(filter_lbl)

                def make_cat_filter(cat_name, frames_list):
                    def do_filter():
                        for frm, cats in frames_list:
                            frm.setVisible(cat_name in cats)
                    return do_filter

                def make_show_all(frames_list):
                    def do_show():
                        for frm, cats in frames_list:
                            frm.setVisible(True)
                    return do_show

                show_all_btn = QPushButton("Show All")
                show_all_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                show_all_btn.setStyleSheet(f"""
                    QPushButton {{
                        font-size: 11px; font-weight: 600; color: #1f2937;
                        background-color: #e5e7eb; border: 1px solid #d1d5db;
                        border-radius: 3px; padding: 3px 8px;
                    }}
                    QPushButton:hover {{ background-color: #d1d5db; }}
                """)
                show_all_btn.clicked.connect(make_show_all(filterable_entry_frames))
                filter_layout.addWidget(show_all_btn)

                for cat in sorted(all_cats):
                    cat_btn = QPushButton(cat)
                    cat_btn.setCursor(Qt.CursorShape.PointingHandCursor)
                    cat_btn.setStyleSheet(f"""
                        QPushButton {{
                            font-size: 11px; font-weight: 600; color: {accent_color};
                            background-color: {bg_color}; border: 1px solid {accent_color}66;
                            border-radius: 3px; padding: 3px 8px;
                        }}
                        QPushButton:hover {{ background-color: {accent_color}22; }}
                    """)
                    cat_btn.clicked.connect(make_cat_filter(cat, filterable_entry_frames))
                    filter_layout.addWidget(cat_btn)

                filter_layout.addStretch()
                incidents_layout.addWidget(filter_frame)

        for match in sorted_matches:
            date = match['date']
            text = match['text']
            matched_terms = match.get('matched_terms', [])
            categories = match.get('categories', [])

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Create highlighted HTML
            escaped_text = html.escape(text)
            highlighted_html = escaped_text
            for term in matched_terms[:5]:
                try:
                    escaped_term = html.escape(term)
                    pattern = re.compile(re.escape(escaped_term), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_term}</span>',
                        highlighted_html
                    )
                except:
                    pass
            highlighted_html = highlighted_html.replace('\n', '<br>')

            # Entry frame with left border
            entry_frame = QFrame()
            entry_frame.setObjectName("allNotesEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#allNotesEntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {accent_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button
            toggle_btn = QPushButton("\u25b8")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            # Date label
            date_label = QLabel(f"{date_str}")
            date_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 16px;
                    font-weight: 600;
                    color: {accent_color};
                    background: transparent;
                    border: none;
                }}
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Category labels (e.g., "Childhood Trauma", "PTSD Related")
            for cat_label in categories[:2]:
                cat_badge = QLabel(cat_label)
                cat_badge.setStyleSheet(f"""
                    QLabel {{
                        font-size: 11px;
                        font-weight: 700;
                        color: #1f2937;
                        background-color: {accent_color}30;
                        border: 1px solid {accent_color}66;
                        border-radius: 3px;
                        padding: 2px 6px;
                    }}
                """)
                header_row.addWidget(cat_badge)

            # Matched terms badges (first 3) - dark text for readability
            for term in matched_terms[:3]:
                term_badge = QLabel(term)
                term_badge.setStyleSheet(f"""
                    QLabel {{
                        font-size: 11px;
                        font-weight: 600;
                        color: #1f2937;
                        background-color: {accent_color}20;
                        border: 1px solid {accent_color}55;
                        border-radius: 3px;
                        padding: 2px 6px;
                    }}
                """)
                header_row.addWidget(term_badge)

            header_row.addStretch()

            # Checkbox
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body text (hidden by default)
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"""
                QTextEdit {{
                    font-size: 15px;
                    color: #333;
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }}
            """)
            body_text.setVisible(False)

            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))

            # Toggle function
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("\u25b8")
                    else:
                        body.setVisible(True)
                        btn.setText("\u25be")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout.addWidget(body_text)
            incidents_layout.addWidget(entry_frame)

            # Track for category filtering (H8, H10)
            if key in ('h8', 'h10'):
                filterable_entry_frames.append((entry_frame, categories))

        # Store checkboxes reference
        setattr(self, f"popup_{key}_imported_checkboxes", checkboxes)

        print(f"[HCR-20] Displayed {len(sorted_matches)} entries for {key.upper()}")
        return len(sorted_matches) > 0

    def _populate_imported_entries_forensic(self, key: str, entries_layout) -> bool:
        """
        Populate H1/H2 imported data section by searching ALL notes with broad keyword patterns.
        Similar to how H3/H4 work - directly searches notes for violence/antisocial keywords.

        Args:
            key: The item key ('h1' for Violence, 'h2' for Antisocial Behaviour)
            entries_layout: The QVBoxLayout to add entries to

        Returns:
            True if any entries were added
        """
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy, QScrollArea, QWidget
        from PySide6.QtCore import Qt
        import html
        import re

        # Get raw notes
        notes = getattr(self, '_extracted_raw_notes', [])
        if not notes:
            print(f"[HCR-20 {key.upper()}] No raw notes available")
            return False

        # Log date range of notes being searched
        dates = []
        for note in notes:
            if isinstance(note, dict):
                dt = note.get('date') or note.get('datetime')
                if dt:
                    dates.append(dt)
        if dates:
            try:
                min_date = min(dates)
                max_date = max(dates)
                print(f"[HCR-20 {key.upper()}] Searching ALL {len(notes)} raw notes (date range: {min_date} to {max_date})")
            except:
                print(f"[HCR-20 {key.upper()}] Searching ALL {len(notes)} raw notes")
        else:
            print(f"[HCR-20 {key.upper()}] Searching ALL {len(notes)} raw notes (no dates found)")

        # Clear existing entries
        checkboxes = getattr(self, f"popup_{key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            item = entries_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Define broad keyword patterns for H1 (Violence) and H2 (Antisocial Behaviour)
        # These are word-boundary patterns to search ALL notes directly
        if key == 'h1':
            # H1: Violence - physical aggression, assault, harm, weapons
            PATTERNS = [
                # Physical violence
                r'\bviolence\b', r'\bviolent\b', r'\bviolently\b',
                r'\bassault\b', r'\bassaulted\b', r'\bassaulting\b', r'\bassaults\b',
                r'\battack\b', r'\battacked\b', r'\battacking\b', r'\battacks\b',
                r'\bpunch\b', r'\bpunched\b', r'\bpunching\b', r'\bpunches\b',
                r'\bkick\b', r'\bkicked\b', r'\bkicking\b', r'\bkicks\b',
                r'\bslap\b', r'\bslapped\b', r'\bslapping\b', r'\bslaps\b',
                r'\bhit\b', r'\bhitting\b', r'\bhits\b',
                r'\bstruck\b', r'\bstriking\b', r'\bstrikes\b', r'\bstrike\b',
                r'\bbeat\b', r'\bbeating\b', r'\bbeaten\b',
                r'\bfight\b', r'\bfighting\b', r'\bfights\b', r'\bfought\b',
                r'\bbite\b', r'\bbit\b', r'\bbitten\b', r'\bbiting\b',
                r'\bscratch\b', r'\bscratched\b', r'\bscratching\b',
                r'\bspit\b', r'\bspat\b', r'\bspitting\b',
                r'\bhead.?butt\b', r'\bheadbutt\b', r'\bhead-butt\b',
                # Aggression
                r'\baggression\b', r'\baggressive\b', r'\baggressively\b',
                r'\bphysically aggressive\b', r'\bphysical aggression\b',
                r'\blashed out\b', r'\blashing out\b',
                r'\bphysical altercation\b',
                # Restraint
                r'\brestraint\b', r'\brestrained\b', r'\brestraining\b',
                r'\bprone restraint\b', r'\bsupine restraint\b',
                r'\bphysical intervention\b', r'\bbreakaway\b',
                r'\bcontrol and restraint\b', r'\bc&r\b',
                r'\brapid tranquil\b', r'\brt\b',
                # Harm
                r'\bharm\b', r'\bharmed\b', r'\bharming\b', r'\bharmful\b',
                r'\binjury\b', r'\binjured\b', r'\binjuring\b', r'\binjuries\b',
                r'\bhurt\b', r'\bhurting\b',
                r'\bwound\b', r'\bwounded\b', r'\bwounding\b',
                # Weapons/objects
                r'\bweapon\b', r'\bweapons\b', r'\barmed\b',
                r'\bknife\b', r'\bknives\b', r'\bblade\b', r'\bsharp object\b',
                r'\bthrew\b', r'\bthrown\b', r'\bthrowing\b',
                # Murder/homicide
                r'\bmurder\b', r'\bmurdered\b', r'\bmurdering\b',
                r'\bhomicide\b', r'\bmanslaughter\b',
                r'\bkilled\b', r'\bkilling\b', r'\bkill\b',
                # Strangulation
                r'\bstrangle\b', r'\bstrangled\b', r'\bstrangling\b', r'\bstrangulation\b',
                r'\bchoke\b', r'\bchoked\b', r'\bchoking\b',
                # Threatening
                r'\bthreatened\b', r'\bthreatening\b', r'\bthreaten\b', r'\bthreats\b',
                r'\bmenacing\b', r'\bintimidating\b', r'\bintimidation\b',
                # Index offence related
                r'\bindex offence\b', r'\bindex offense\b',
                r'\bconviction\b', r'\bconvicted\b', r'\bconvictions\b',
                r'\boffence\b', r'\boffences\b', r'\boffense\b', r'\boffenses\b',
                r'\bsection 37\b', r'\bsection 41\b', r'\bs\.37\b', r'\bs\.41\b',
                r'\bdiminished responsibility\b',
                r'\bgbh\b', r'\babh\b', r'\bactual bodily harm\b', r'\bgrievous bodily harm\b',
            ]
            CATEGORY_LABEL = "Violence"
            CATEGORY_COLOR = "#b71c1c"  # Dark red
        else:
            # H2: Antisocial Behaviour - property damage, verbal aggression, threats, bullying
            PATTERNS = [
                # Property damage
                r'\bdamage\b', r'\bdamaged\b', r'\bdamaging\b', r'\bdamages\b',
                r'\bdestroy\b', r'\bdestroyed\b', r'\bdestroying\b', r'\bdestruction\b',
                r'\bsmash\b', r'\bsmashed\b', r'\bsmashing\b',
                r'\bbreak\b', r'\bbroke\b', r'\bbroken\b', r'\bbreaking\b',
                r'\bvandal\b', r'\bvandalism\b', r'\bvandalised\b', r'\bvandalized\b',
                r'\bkicking\s+(door|wall|furniture)\b',
                r'\bpunching\s+(door|wall|furniture)\b',
                r'\boverturned\b', r'\bupended\b', r'\bflipped\b',
                # Verbal aggression
                r'\bverbally aggressive\b', r'\bverbal aggression\b',
                r'\bverbally abusive\b', r'\bverbal abuse\b',
                r'\bshouting\b', r'\bshouted\b', r'\byelling\b', r'\byelled\b',
                r'\bscreaming\b', r'\bscreamed\b',
                r'\brabusive language\b', r'\babusive\b',
                r'\bswearing\b', r'\bswore\b', r'\bsworn\b',
                r'\bname.?calling\b', r'\bcalled.{0,20}names\b',
                r'\binsulting\b', r'\binsulted\b', r'\binsults\b',
                r'\bhostile\b', r'\bhostility\b',
                r'\bconfrontation\b', r'\bconfrontational\b',
                r'\bargumentative\b', r'\barguing\b', r'\bargument\b',
                # Bullying/exploitation
                r'\bbully\b', r'\bbullying\b', r'\bbullied\b',
                r'\bpicking on\b', r'\bpicked on\b',
                r'\bexploiting\b', r'\bexploited\b', r'\bexploitation\b',
                r'\bcoercion\b', r'\bcoerced\b', r'\bcoercing\b',
                r'\bpressuring\b', r'\bpressured\b',
                r'\btargeting\b', r'\btargeted\b',
                # Antisocial behaviour
                r'\bantisocial\b', r'\banti-social\b',
                r'\bdisruptive\b', r'\bdisruption\b',
                r'\bdisobedient\b', r'\bdisobedience\b',
                r'\buncooperative\b', r'\bnon-compliant\b', r'\bnoncompliant\b',
                r'\bdefiant\b', r'\bdefiance\b',
                r'\bagitated\b', r'\bagitation\b',
                r'\bthrowing objects\b', r'\bthrew objects\b',
                # Criminal behaviour
                r'\btheft\b', r'\bstole\b', r'\bstealing\b', r'\bstolen\b',
                r'\barson\b', r'\bfire.?setting\b',
                r'\bburglary\b', r'\brobbery\b',
                r'\bcriminal damage\b',
                # AWOL/absconding
                r'\bawol\b', r'\babscond\b', r'\babsconded\b', r'\babsconding\b',
                r'\bfailed to return\b', r'\bescape\b', r'\bescaped\b',
                r'\bleft without permission\b', r'\bUAL\b',
            ]
            CATEGORY_LABEL = "Antisocial Behaviour"
            CATEGORY_COLOR = "#ff9800"  # Orange

        def matches_any_pattern(text_lower, patterns):
            """Check if text matches any of the word-boundary patterns."""
            for pattern in patterns:
                if re.search(pattern, text_lower, re.IGNORECASE):
                    return True
            return False

        def get_matched_terms(text_lower, patterns):
            """Get all matching terms from text."""
            matched = []
            for pattern in patterns:
                matches = re.findall(pattern, text_lower, re.IGNORECASE)
                matched.extend(matches)
            return list(set(matched))

        # Search all notes for matches
        all_matches = []
        for note in notes:
            if isinstance(note, dict):
                text = note.get('text', '') or note.get('content', '') or ''
            else:
                text = str(note)

            if not text.strip():
                continue

            text_lower = text.lower()
            if matches_any_pattern(text_lower, PATTERNS):
                note_date = note.get('date') or note.get('datetime') if isinstance(note, dict) else None
                matched_terms = get_matched_terms(text_lower, PATTERNS)
                all_matches.append({
                    'text': text,
                    'date': note_date,
                    'matched_terms': matched_terms,
                    'category': CATEGORY_LABEL,
                })

        if not all_matches:
            print(f"[HCR-20 {key.upper()}] No matches found in {len(notes)} notes")
            return False

        # Sort by date (newest first)
        def get_sort_date(x):
            d = x.get("date")
            if d is None:
                from datetime import datetime
                return datetime.min
            return d

        sorted_matches = sorted(all_matches, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep only the most relevant entry per date
        seen_dates = {}
        deduplicated_matches = []
        for match in sorted_matches:
            d = match.get("date")
            if d is None:
                date_key = ""
            elif hasattr(d, "strftime"):
                date_key = d.strftime("%Y-%m-%d")
            else:
                date_key = str(d)

            # Calculate relevance score: number of matched terms + text length
            num_terms = len(match.get('matched_terms', []))
            text_len = len(match.get('text', ''))
            relevance_score = num_terms * 100 + text_len

            if date_key not in seen_dates:
                seen_dates[date_key] = (match, relevance_score)
                deduplicated_matches.append(match)
            else:
                existing_match, existing_score = seen_dates[date_key]
                if relevance_score > existing_score:
                    deduplicated_matches.remove(existing_match)
                    deduplicated_matches.append(match)
                    seen_dates[date_key] = (match, relevance_score)

        sorted_matches = sorted(deduplicated_matches, key=get_sort_date, reverse=True)

        print(f"[HCR-20 {key.upper()}] Found {len(sorted_matches)} matching notes (deduplicated by date)")

        # Store all matches for later use
        setattr(self, f"_forensic_{key}_all_incidents", sorted_matches)
        setattr(self, f"_forensic_{key}_current_filter", None)
        setattr(self, f"_forensic_{key}_cat_colors", {CATEGORY_LABEL: CATEGORY_COLOR})
        setattr(self, f"_forensic_{key}_sev_colors", {"high": "#dc2626", "medium": "#f59e0b", "low": "#22c55e"})

        # Create header with count
        header_widget = QWidget()
        header_widget.setStyleSheet("background: transparent;")
        header_layout = QHBoxLayout(header_widget)
        header_layout.setContentsMargins(0, 0, 0, 8)
        header_layout.setSpacing(8)

        count_badge = QLabel(f"{len(sorted_matches)} notes found")
        count_badge.setStyleSheet(f"""
            QLabel {{
                font-size: 13px;
                font-weight: 600;
                color: white;
                background: {CATEGORY_COLOR};
                border: none;
                border-radius: 4px;
                padding: 4px 10px;
            }}
        """)
        header_layout.addWidget(count_badge)
        header_layout.addStretch()
        entries_layout.addWidget(header_widget)

        # Container for entries (for re-rendering)
        incidents_container = QWidget()
        incidents_container.setObjectName(f"incidentsContainer_{key}")
        incidents_container.setStyleSheet("background: transparent;")
        incidents_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Minimum)
        incidents_layout = QVBoxLayout(incidents_container)
        incidents_layout.setContentsMargins(0, 0, 0, 0)
        incidents_layout.setSpacing(8)
        entries_layout.addWidget(incidents_container)

        # Store layout reference
        setattr(self, f"_forensic_{key}_incidents_layout", incidents_layout)

        # Render ALL matches (no limit - matching H3 behaviour)
        self._render_forensic_incidents_direct(key, sorted_matches, checkboxes, CATEGORY_COLOR)

        # Store checkboxes reference
        setattr(self, f"popup_{key}_imported_checkboxes", checkboxes)

        print(f"[HCR-20] Displayed {len(sorted_matches)} entries for {key.upper()}")
        return len(sorted_matches) > 0

    def _render_forensic_incidents_direct(self, key: str, matches: list, checkboxes: list, category_color: str):
        """Render the list of direct keyword matches for H1/H2.
        Matching H3 _populate_h3_subsection style: toggle button, date, checkbox, expandable body."""
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QLabel, QCheckBox, QTextEdit, QPushButton, QSizePolicy
        from PySide6.QtCore import Qt
        import html
        import re

        incidents_layout = getattr(self, f"_forensic_{key}_incidents_layout", None)
        if not incidents_layout:
            return

        # Clear existing
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while incidents_layout.count():
            child = incidents_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()

        # Use category color as accent (dark red for H1, orange for H2)
        accent_color = category_color
        bg_color = "rgba(255, 230, 230, 0.95)" if key == 'h1' else "rgba(255, 243, 224, 0.95)"
        border_color = f"{category_color}66"

        for match in matches:
            date = match["date"]
            text = match["text"]
            matched_terms = match.get("matched_terms", [])

            if not text or not text.strip():
                continue

            # Format date
            if date:
                if hasattr(date, "strftime"):
                    date_str = date.strftime("%d %b %Y")
                else:
                    date_str = str(date)
            else:
                date_str = "No date"

            # Create HTML with highlighted matched terms
            escaped_text = html.escape(text)
            highlighted_html = escaped_text
            for term in matched_terms:
                try:
                    escaped_term = html.escape(term)
                    pattern = re.compile(re.escape(escaped_term), re.IGNORECASE)
                    highlighted_html = pattern.sub(
                        f'<span style="background-color: #FFEB3B; color: #000; font-weight: bold; padding: 1px 3px; border-radius: 3px;">{escaped_term}</span>',
                        highlighted_html
                    )
                except:
                    pass
            highlighted_html = highlighted_html.replace('\n', '<br>')

            # Create entry frame with colored left border (matching H3 style)
            entry_frame = QFrame()
            entry_frame.setObjectName("forensicEntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#forensicEntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {accent_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row with toggle, date, term badges, and checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            # Toggle button (matching H3)
            toggle_btn = QPushButton("\u25b8")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            # Date label (matching H3)
            date_label = QLabel(f"{date_str}")
            date_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 16px;
                    font-weight: 600;
                    color: {accent_color};
                    background: transparent;
                    border: none;
                }}
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            # Matched terms badges (show first 3)
            for term in matched_terms[:3]:
                term_badge = QLabel(term)
                term_badge.setStyleSheet(f"""
                    QLabel {{
                        font-size: 11px;
                        font-weight: 600;
                        color: white;
                        background: {accent_color}cc;
                        border: none;
                        border-radius: 3px;
                        padding: 2px 6px;
                    }}
                """)
                header_row.addWidget(term_badge)

            header_row.addStretch()

            # Checkbox (matching H3)
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)
            cb.stateChanged.connect(lambda _, k=key: self._update_preview(k))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body text (hidden by default, with highlighted matched terms - matching H3)
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"""
                QTextEdit {{
                    font-size: 15px;
                    color: #333;
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }}
            """)
            body_text.setVisible(False)

            # Calculate height based on content
            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))

            # Toggle function (matching H3)
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("\u25b8")
                    else:
                        body.setVisible(True)
                        btn.setText("\u25be")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout.addWidget(body_text)
            incidents_layout.addWidget(entry_frame)

    def _populate_h3_imported_entries(self, extraction_result: dict) -> bool:
        """
        Populate H3 (Relationships) with separate imported data sections for
        Intimate and Non-intimate relationships.

        Searches ALL raw notes directly (not using extractor's filtered results)
        to ensure complete coverage.

        Args:
            extraction_result: The extraction result from HCR20Extractor (not used - we search raw notes)

        Returns:
            True if any entries were added
        """
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy
        from PySide6.QtCore import Qt
        import html
        import re

        populated = False
        key = 'h3'

        # Get ALL raw notes directly - bypass extractor filtering
        all_notes = getattr(self, '_extracted_raw_notes', [])
        if not all_notes:
            print("[HCR-20 H3] No raw notes available")
            return False

        # Log date range of notes being searched
        dates = []
        for note in all_notes:
            if isinstance(note, dict):
                dt = note.get('date') or note.get('datetime')
                if dt:
                    dates.append(dt)
        if dates:
            try:
                min_date = min(dates)
                max_date = max(dates)
                print(f"[HCR-20 H3] Searching ALL {len(all_notes)} raw notes (date range: {min_date} to {max_date})")
            except:
                print(f"[HCR-20 H3] Searching ALL {len(all_notes)} raw notes")
        else:
            print(f"[HCR-20 H3] Searching ALL {len(all_notes)} raw notes (no dates found)")

        # Keywords to categorize matches - using word boundary patterns to avoid cross-matching
        # INTIMATE: partner relationships (romantic/sexual)
        # These patterns are checked FIRST and take priority
        INTIMATE_PATTERNS = [
            r'\bpartner\b', r'\bpartners\b', r'\bwife\b', r'\bhusband\b',
            r'\bgirlfriend\b', r'\bboyfriend\b', r'\bspouse\b',
            r'\bmarried\b', r'\bmarriage\b', r'\bdivorced\b', r'\bdivorce\b',
            r'\bseparation\b', r'\bseparated\b',
            r'\bdomestic violence\b', r'\bdomestic abuse\b', r'\bdomestic relationship\b',
            r'\brelationship breakdown\b', r'\bex-partner\b', r'\bex partner\b',
            r'\bintimate\b', r'\bromantic\b', r'\bcohabit\b', r'\bcohabiting\b',
            r'\bliving together\b', r'\bengaged to\b', r'\bfiancé\b', r'\bfiancee\b', r'\bfiance\b',
            r'\bex-wife\b', r'\bex-husband\b', r'\bex wife\b', r'\bex husband\b',
            r'\bex-girlfriend\b', r'\bex-boyfriend\b', r'\bex girlfriend\b', r'\bex boyfriend\b',
        ]

        # NON-INTIMATE: family, friends, social relationships
        # Careful to avoid matching parts of intimate words
        NON_INTIMATE_PATTERNS = [
            r'\bfamily\b', r'\bfamilies\b', r'\bfamilial\b',
            r'\bmother\b', r'\bfather\b', r'\bparent\b', r'\bparents\b', r'\bparental\b',
            r'\bsibling\b', r'\bsiblings\b', r'\bbrother\b', r'\bsister\b',
            r'\bfriend\b', r'\bfriends\b', r'\bfriendship\b', r'\bfriendships\b',
            r'\bcolleague\b', r'\bcolleagues\b', r'\bpeer\b', r'\bpeers\b',
            r'\bneighbour\b', r'\bneighbor\b', r'\bneighbours\b', r'\bneighbors\b',
            r'\bsocially isolated\b', r'\bsocial isolation\b', r'\bsupport network\b',
            r'\bchildren\b', r'\bchild\b', r'\bson\b', r'\bdaughter\b', r'\bsons\b', r'\bdaughters\b',
            r'\baunt\b', r'\buncle\b', r'\bcousin\b', r'\bcousins\b',
            r'\bgrandparent\b', r'\bgrandparents\b', r'\bgrandmother\b', r'\bgrandfather\b',
            r'\bgrandson\b', r'\bgranddaughter\b', r'\bnephew\b', r'\bniece\b',
            r'\bin-law\b', r'\bin law\b', r'\bstepmother\b', r'\bstepfather\b',
            r'\bstepson\b', r'\bstepdaughter\b', r'\bhalf-brother\b', r'\bhalf-sister\b',
        ]

        def matches_any_pattern(text_lower, patterns):
            """Check if text matches any of the word-boundary patterns."""
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    return True
            return False

        def get_matched_terms(text_lower, patterns):
            """Get list of matched terms from patterns."""
            matched = []
            for pattern in patterns:
                match = re.search(pattern, text_lower)
                if match:
                    matched.append(match.group())
            return matched

        # Categorize ALL notes by keyword - search every note
        intimate_matches = []
        non_intimate_matches = []

        for note in all_notes:
            # Get note text
            if isinstance(note, dict):
                text = note.get('text', '') or note.get('content', '') or note.get('body', '') or ''
            else:
                text = str(note) if note else ''

            if not text.strip():
                continue

            text_lower = text.lower()

            # Check for keywords using word boundaries
            is_intimate = matches_any_pattern(text_lower, INTIMATE_PATTERNS)
            is_non_intimate = matches_any_pattern(text_lower, NON_INTIMATE_PATTERNS)

            # Build match entry
            note_date = note.get('date') or note.get('datetime') if isinstance(note, dict) else None

            if is_intimate:
                # Get the matched intimate terms for highlighting
                matched_terms = get_matched_terms(text_lower, INTIMATE_PATTERNS)
                intimate_matches.append({
                    'note': note,
                    'date': note_date,
                    'matches': [{'term': t, 'excerpt': ''} for t in matched_terms],
                })

            if is_non_intimate:
                # Get the matched non-intimate terms for highlighting
                matched_terms = get_matched_terms(text_lower, NON_INTIMATE_PATTERNS)
                non_intimate_matches.append({
                    'note': note,
                    'date': note_date,
                    'matches': [{'term': t, 'excerpt': ''} for t in matched_terms],
                })

        print(f"[HCR-20 H3] Found {len(intimate_matches)} intimate matches, {len(non_intimate_matches)} non-intimate matches")

        # Populate Intimate Relationships section
        intimate_populated = self._populate_h3_subsection(
            'intimate', intimate_matches,
            bg_color="rgba(255, 230, 230, 0.95)",
            border_color="rgba(180, 100, 100, 0.5)",
            accent_color="#8B0000"
        )
        if intimate_populated:
            populated = True
            # Expand the section
            intimate_section = getattr(self, "popup_h3_intimate_imported_section", None)
            if intimate_section and hasattr(intimate_section, '_is_collapsed'):
                if intimate_section._is_collapsed:
                    intimate_section._toggle_collapse()

        # Populate Non-intimate Relationships section
        non_intimate_populated = self._populate_h3_subsection(
            'non_intimate', non_intimate_matches,
            bg_color="rgba(230, 255, 230, 0.95)",
            border_color="rgba(100, 180, 100, 0.5)",
            accent_color="#006400"
        )
        if non_intimate_populated:
            populated = True
            # Expand the section
            non_intimate_section = getattr(self, "popup_h3_non_intimate_imported_section", None)
            if non_intimate_section and hasattr(non_intimate_section, '_is_collapsed'):
                if non_intimate_section._is_collapsed:
                    non_intimate_section._toggle_collapse()

        return populated

    def _populate_h3_subsection(self, section_key: str, matches: list,
                                 bg_color: str, border_color: str, accent_color: str) -> bool:
        """
        Populate one of the H3 imported data subsections (intimate or non_intimate).

        When checkboxes are ticked, the text is added to the corresponding subsection text field.
        """
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy
        from PySide6.QtCore import Qt
        import html
        import re

        key = 'h3'
        entries_layout = getattr(self, f"popup_{key}_{section_key}_imported_entries_layout", None)
        if not entries_layout:
            return False

        # Clear existing entries
        checkboxes = getattr(self, f"popup_{key}_{section_key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            item = entries_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not matches:
            return False

        # Sort by date (newest first)
        def get_sort_date(match):
            dt = match.get('date')
            if dt is None:
                return ""
            if hasattr(dt, "strftime"):
                return dt.strftime("%Y-%m-%d")
            return str(dt)

        sorted_matches = sorted(matches, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep only the most relevant entry per date
        seen_dates = {}
        deduplicated_matches = []
        for match in sorted_matches:
            date_key = get_sort_date(match)
            num_matches = len(match.get('matches', []))
            note = match.get('note', {})
            text_len = len(note.get('text', '') or note.get('content', '') or '') if isinstance(note, dict) else 0
            relevance_score = num_matches * 100 + text_len

            if date_key not in seen_dates:
                seen_dates[date_key] = (match, relevance_score)
                deduplicated_matches.append(match)
            else:
                existing_match, existing_score = seen_dates[date_key]
                if relevance_score > existing_score:
                    deduplicated_matches.remove(existing_match)
                    deduplicated_matches.append(match)
                    seen_dates[date_key] = (match, relevance_score)

        sorted_matches = sorted(deduplicated_matches, key=get_sort_date, reverse=True)

        # Get reference to the text widget for this subsection
        text_widget = getattr(self, f"popup_{key}_{section_key}_text_widget", None)

        # No limit for historical items - show ALL entries (deduplicated by date)
        entries_added = 0
        for match in sorted_matches:
            dt = match.get('date')
            note = match.get('note', {})
            source = note.get('source', '') if isinstance(note, dict) else ''

            # Get the FULL note content
            if isinstance(note, dict):
                text = note.get('text', '') or note.get('content', '') or ''
            else:
                text = str(note) if note else ''

            # If no full text, fall back to matched excerpts
            if not text.strip():
                match_texts = []
                for term_match in match.get('matches', []):
                    excerpt = term_match.get('excerpt', '').strip()
                    if excerpt:
                        match_texts.append(excerpt)
                text = "\n\n".join(match_texts)

            if not text.strip():
                continue

            # Collect matched terms for highlighting
            matched_terms = set()
            for term_match in match.get('matches', []):
                term = term_match.get('term', '')
                if term:
                    matched_terms.add(term.lower())

            # Create highlighted HTML
            escaped_text = html.escape(text)
            highlighted_text = escaped_text
            for term in matched_terms:
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                highlighted_text = pattern.sub(
                    lambda m: f'<span style="background-color: #ffeb3b; padding: 1px 3px; border-radius: 3px; font-weight: 600;">{m.group()}</span>',
                    highlighted_text
                )
            highlighted_html = highlighted_text.replace('\n', '<br>')

            # Format date
            if dt:
                if hasattr(dt, "strftime"):
                    date_str = dt.strftime("%d %b %Y")
                else:
                    date_str = str(dt)
            else:
                date_str = "No date"

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("h3EntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#h3EntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {accent_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row with toggle, date, source, and checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            date_label = QLabel(f"{date_str}")
            date_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 16px;
                    font-weight: 600;
                    color: {accent_color};
                    background: transparent;
                    border: none;
                }}
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            if source:
                source_label = QLabel(f"({source})")
                source_label.setStyleSheet("""
                    QLabel {
                        font-size: 14px;
                        color: #666;
                        background: transparent;
                        border: none;
                    }
                """)
                header_row.addWidget(source_label)

            header_row.addStretch()

            # Checkbox - when ticked, adds text to the subsection text field
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setProperty("section_key", section_key)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)

            # Connect checkbox to update both preview and subsection text field
            def make_checkbox_handler(checkbox, text_w, sk):
                def handler(state):
                    self._update_h3_subsection_from_checkboxes(sk)
                    self._update_preview('h3')
                return handler

            cb.stateChanged.connect(make_checkbox_handler(cb, text_widget, section_key))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body text (hidden by default) - with highlighted matched terms
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"""
                QTextEdit {{
                    font-size: 15px;
                    color: #333;
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }}
            """)
            body_text.setVisible(False)

            # Calculate height based on content
            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))  # Reduced by 10% (was 160)

            # Toggle function
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("▸")
                    else:
                        body.setVisible(True)
                        btn.setText("▾")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout.addWidget(body_text)
            entries_layout.addWidget(entry_frame)
            entries_added += 1

        # Store checkboxes reference
        setattr(self, f"popup_{key}_{section_key}_imported_checkboxes", checkboxes)

        return entries_added > 0

    def _update_h3_subsection_from_checkboxes(self, section_key: str):
        """
        Update the H3 subsection text field based on checked imported data checkboxes.
        """
        key = 'h3'
        checkboxes = getattr(self, f"popup_{key}_{section_key}_imported_checkboxes", [])
        text_widget = getattr(self, f"popup_{key}_{section_key}_text_widget", None)

        if not text_widget:
            return

        # Collect text from checked boxes
        checked_texts = []
        for cb in checkboxes:
            if cb.isChecked():
                full_text = cb.property("full_text")
                if full_text:
                    # Get just the first 500 chars as a summary
                    summary = full_text[:500].strip()
                    if len(full_text) > 500:
                        summary += "..."
                    checked_texts.append(summary)

        # Update the text widget with checked content
        if checked_texts:
            text_widget.setPlainText("\n\n---\n\n".join(checked_texts))
        else:
            text_widget.clear()

    def _populate_h4_imported_entries(self, extraction_result: dict) -> bool:
        """
        Populate H4 (Employment) with separate imported data sections for
        Education and Employment.

        Searches ALL raw notes directly to ensure complete coverage.

        Args:
            extraction_result: The extraction result from HCR20Extractor (not used - we search raw notes)

        Returns:
            True if any entries were added
        """
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy
        from PySide6.QtCore import Qt
        import html
        import re

        populated = False
        key = 'h4'

        # Get ALL raw notes directly - bypass extractor filtering
        all_notes = getattr(self, '_extracted_raw_notes', [])
        if not all_notes:
            print("[HCR-20 H4] No raw notes available")
            return False

        # Log date range of notes
        dates = []
        for note in all_notes:
            if isinstance(note, dict):
                dt = note.get('date') or note.get('datetime')
                if dt:
                    dates.append(dt)
        if dates:
            try:
                min_date = min(dates)
                max_date = max(dates)
                print(f"[HCR-20 H4] Searching ALL {len(all_notes)} raw notes (date range: {min_date} to {max_date})")
            except:
                print(f"[HCR-20 H4] Searching ALL {len(all_notes)} raw notes")
        else:
            print(f"[HCR-20 H4] Searching ALL {len(all_notes)} raw notes (no dates found)")

        # Keywords to categorize matches - using word boundary patterns
        # EDUCATION: schools, colleges, universities, qualifications
        EDUCATION_PATTERNS = [
            r'\bschool\b', r'\bschools\b', r'\bprimary school\b', r'\bsecondary school\b',
            r'\bcollege\b', r'\buniversity\b', r'\beducation\b', r'\beducational\b',
            r'\bqualification\b', r'\bqualifications\b', r'\bgcse\b', r'\ba-level\b', r'\ba level\b',
            r'\bdegree\b', r'\bgraduate\b', r'\bgraduated\b', r'\bstudent\b', r'\bstudying\b',
            r'\blearning\b', r'\blearning difficult\b', r'\blearning disabil\b',
            r'\bexam\b', r'\bexams\b', r'\bexpelled\b', r'\bexpulsion\b',
            r'\bsuspended\b', r'\bsuspension\b', r'\btruancy\b', r'\btruant\b',
            r'\bspecial educational needs\b', r'\bsen\b', r'\behcp\b',
            r'\bacademic\b', r'\bteacher\b', r'\bclassroom\b', r'\bpupil\b',
        ]

        # EMPLOYMENT: jobs, work, occupations, vocational
        # Note: Exclude 'working on' followed by therapeutic terms
        EMPLOYMENT_PATTERNS = [
            r'\bemployed\b', r'\bemployment\b', r'\bunemployed\b', r'\bunemployment\b',
            r'\bjob\b', r'\bjobs\b', r'\boccupation\b', r'\bvocational\b',
            r'\bcareer\b', r'\bprofession\b', r'\bprofessional\b',
            r'\bdismissed\b', r'\bsacked\b', r'\bfired\b', r'\bredundant\b', r'\bredundancy\b',
            r'\bresigned\b', r'\bretired\b', r'\bretirement\b',
            r'\bworkplace\b', r'\bemployer\b', r'\bemployers\b',
            r'\bself-employed\b', r'\bself employed\b',
            r'\bfull-time\b', r'\bfull time\b', r'\bpart-time\b', r'\bpart time\b',
            r'\bbenefits\b', r'\bjobseeker\b', r'\bjsa\b', r'\besa\b', r'\buniversal credit\b',
            r'\bvocational training\b', r'\bapprentice\b', r'\bapprenticeship\b',
            r'\bwork experience\b', r'\bwork placement\b',
            # 'working' only when followed by employment context
            r'\bworking at\b', r'\bworking for\b', r'\bworking in\b', r'\bworking as\b',
            r'\bnot working\b', r'\bstopped working\b', r'\bstarted working\b',
            r'\bworked at\b', r'\bworked for\b', r'\bworked in\b', r'\bworked as\b',
        ]

        def matches_any_pattern(text_lower, patterns):
            """Check if text matches any of the word-boundary patterns."""
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    return True
            return False

        def get_matched_terms(text_lower, patterns):
            """Get list of matched terms from patterns."""
            matched = []
            for pattern in patterns:
                match = re.search(pattern, text_lower)
                if match:
                    matched.append(match.group())
            return matched

        # Categorize ALL notes by keyword
        education_matches = []
        employment_matches = []

        for note in all_notes:
            # Get note text
            if isinstance(note, dict):
                text = note.get('text', '') or note.get('content', '') or note.get('body', '') or ''
            else:
                text = str(note) if note else ''

            if not text.strip():
                continue

            text_lower = text.lower()

            # Check for keywords using word boundaries
            is_education = matches_any_pattern(text_lower, EDUCATION_PATTERNS)
            is_employment = matches_any_pattern(text_lower, EMPLOYMENT_PATTERNS)

            # Build match entry
            note_date = note.get('date') or note.get('datetime') if isinstance(note, dict) else None

            if is_education:
                matched_terms = get_matched_terms(text_lower, EDUCATION_PATTERNS)
                education_matches.append({
                    'note': note,
                    'date': note_date,
                    'matches': [{'term': t, 'excerpt': ''} for t in matched_terms],
                })

            if is_employment:
                matched_terms = get_matched_terms(text_lower, EMPLOYMENT_PATTERNS)
                employment_matches.append({
                    'note': note,
                    'date': note_date,
                    'matches': [{'term': t, 'excerpt': ''} for t in matched_terms],
                })

        print(f"[HCR-20 H4] Found {len(education_matches)} education matches, {len(employment_matches)} employment matches")

        # Populate Education section
        education_populated = self._populate_h4_subsection(
            'education', education_matches,
            bg_color="rgba(230, 230, 255, 0.95)",
            border_color="rgba(100, 100, 180, 0.5)",
            accent_color="#00008B"
        )
        if education_populated:
            populated = True
            education_section = getattr(self, "popup_h4_education_imported_section", None)
            if education_section and hasattr(education_section, '_is_collapsed'):
                if education_section._is_collapsed:
                    education_section._toggle_collapse()

        # Populate Employment section
        employment_populated = self._populate_h4_subsection(
            'employment', employment_matches,
            bg_color="rgba(255, 245, 230, 0.95)",
            border_color="rgba(180, 140, 100, 0.5)",
            accent_color="#8B4500"
        )
        if employment_populated:
            populated = True
            employment_section = getattr(self, "popup_h4_employment_imported_section", None)
            if employment_section and hasattr(employment_section, '_is_collapsed'):
                if employment_section._is_collapsed:
                    employment_section._toggle_collapse()

        return populated

    def _populate_h4_subsection(self, section_key: str, matches: list,
                                 bg_color: str, border_color: str, accent_color: str) -> bool:
        """
        Populate one of the H4 imported data subsections (education or employment).

        When checkboxes are ticked, the text is added to the corresponding subsection text field.
        """
        from PySide6.QtWidgets import QFrame, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, QCheckBox, QTextEdit, QSizePolicy
        from PySide6.QtCore import Qt
        import html
        import re

        key = 'h4'
        entries_layout = getattr(self, f"popup_{key}_{section_key}_imported_entries_layout", None)
        if not entries_layout:
            return False

        # Clear existing entries
        checkboxes = getattr(self, f"popup_{key}_{section_key}_imported_checkboxes", [])
        for cb in checkboxes:
            cb.deleteLater()
        checkboxes.clear()

        while entries_layout.count():
            item = entries_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not matches:
            return False

        # Sort by date (newest first)
        def get_sort_date(match):
            dt = match.get('date')
            if dt is None:
                return ""
            if hasattr(dt, "strftime"):
                return dt.strftime("%Y-%m-%d")
            return str(dt)

        sorted_matches = sorted(matches, key=get_sort_date, reverse=True)

        # Deduplicate by date - keep only the most relevant entry per date
        seen_dates = {}
        deduplicated_matches = []
        for match in sorted_matches:
            date_key = get_sort_date(match)
            num_matches = len(match.get('matches', []))
            note = match.get('note', {})
            text_len = len(note.get('text', '') or note.get('content', '') or '') if isinstance(note, dict) else 0
            relevance_score = num_matches * 100 + text_len

            if date_key not in seen_dates:
                seen_dates[date_key] = (match, relevance_score)
                deduplicated_matches.append(match)
            else:
                existing_match, existing_score = seen_dates[date_key]
                if relevance_score > existing_score:
                    deduplicated_matches.remove(existing_match)
                    deduplicated_matches.append(match)
                    seen_dates[date_key] = (match, relevance_score)

        sorted_matches = sorted(deduplicated_matches, key=get_sort_date, reverse=True)

        # Get reference to the text widget for this subsection
        text_widget = getattr(self, f"popup_{key}_{section_key}_text_widget", None)

        # No limit for historical items - show ALL entries (deduplicated by date)
        entries_added = 0
        for match in sorted_matches:
            dt = match.get('date')
            note = match.get('note', {})
            source = note.get('source', '') if isinstance(note, dict) else ''

            # Get the FULL note content
            if isinstance(note, dict):
                text = note.get('text', '') or note.get('content', '') or ''
            else:
                text = str(note) if note else ''

            if not text.strip():
                continue

            # Collect matched terms for highlighting
            matched_terms = set()
            for term_match in match.get('matches', []):
                term = term_match.get('term', '')
                if term:
                    matched_terms.add(term.lower())

            # Create highlighted HTML
            escaped_text = html.escape(text)
            highlighted_text = escaped_text
            for term in matched_terms:
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                highlighted_text = pattern.sub(
                    lambda m: f'<span style="background-color: #ffeb3b; padding: 1px 3px; border-radius: 3px; font-weight: 600;">{m.group()}</span>',
                    highlighted_text
                )
            highlighted_html = highlighted_text.replace('\n', '<br>')

            # Format date
            if dt:
                if hasattr(dt, "strftime"):
                    date_str = dt.strftime("%d %b %Y")
                else:
                    date_str = str(dt)
            else:
                date_str = "No date"

            # Create entry frame
            entry_frame = QFrame()
            entry_frame.setObjectName("h4EntryFrame")
            entry_frame.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Minimum)
            entry_frame.setStyleSheet(f"""
                QFrame#h4EntryFrame {{
                    background: rgba(255, 255, 255, 0.95);
                    border: 1px solid {border_color};
                    border-left: 4px solid {accent_color};
                    border-radius: 8px;
                    padding: 4px;
                }}
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(10, 8, 10, 8)
            entry_layout.setSpacing(6)
            entry_layout.setSizeConstraint(QVBoxLayout.SizeConstraint.SetMinAndMaxSize)

            # Header row with toggle, date, source, and checkbox
            header_row = QHBoxLayout()
            header_row.setSpacing(8)

            toggle_btn = QPushButton("▸")
            toggle_btn.setFixedSize(22, 22)
            toggle_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            toggle_btn.setStyleSheet(f"""
                QPushButton {{
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    font-size: 17px;
                    font-weight: bold;
                    color: {accent_color};
                }}
                QPushButton:hover {{ background: {border_color}; }}
            """)
            header_row.addWidget(toggle_btn)

            date_label = QLabel(f"{date_str}")
            date_label.setStyleSheet(f"""
                QLabel {{
                    font-size: 16px;
                    font-weight: 600;
                    color: {accent_color};
                    background: transparent;
                    border: none;
                }}
            """)
            date_label.setCursor(Qt.CursorShape.PointingHandCursor)
            header_row.addWidget(date_label)

            if source:
                source_label = QLabel(f"({source})")
                source_label.setStyleSheet("""
                    QLabel {
                        font-size: 14px;
                        color: #666;
                        background: transparent;
                        border: none;
                    }
                """)
                header_row.addWidget(source_label)

            header_row.addStretch()

            # Checkbox - when ticked, adds text to the subsection text field
            cb = QCheckBox()
            cb.setProperty("full_text", text)
            cb.setProperty("section_key", section_key)
            cb.setFixedSize(18, 18)
            cb.setStyleSheet("""
                QCheckBox { background: transparent; }
                QCheckBox::indicator { width: 16px; height: 16px; }
            """)

            # Connect checkbox to update both preview and subsection text field
            def make_checkbox_handler(checkbox, text_w, sk):
                def handler(state):
                    self._update_h4_subsection_from_checkboxes(sk)
                    self._update_preview('h4')
                return handler

            cb.stateChanged.connect(make_checkbox_handler(cb, text_widget, section_key))
            header_row.addWidget(cb)
            checkboxes.append(cb)

            entry_layout.addLayout(header_row)

            # Body text (hidden by default) - with highlighted matched terms
            body_text = QTextEdit()
            body_text.setHtml(f'<div style="font-size: 15px; color: #333; line-height: 1.5;">{highlighted_html}</div>')
            body_text.setReadOnly(True)
            body_text.setFrameShape(QFrame.Shape.NoFrame)
            body_text.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
            body_text.setStyleSheet(f"""
                QTextEdit {{
                    font-size: 15px;
                    color: #333;
                    background: {bg_color};
                    border: none;
                    border-radius: 4px;
                    padding: 6px;
                }}
            """)
            body_text.setVisible(False)

            # Calculate height based on content
            doc = body_text.document()
            doc.setTextWidth(body_text.viewport().width() if body_text.viewport().width() > 0 else 350)
            content_height = int(doc.size().height()) + 20
            body_text.setFixedHeight(min(content_height, 144))  # Reduced by 10% (was 160)

            # Toggle function
            def make_toggle(btn, body):
                def toggle():
                    if body.isVisible():
                        body.setVisible(False)
                        btn.setText("▸")
                    else:
                        body.setVisible(True)
                        btn.setText("▾")
                return toggle

            toggle_btn.clicked.connect(make_toggle(toggle_btn, body_text))
            date_label.mousePressEvent = lambda e, b=toggle_btn: b.click()

            entry_layout.addWidget(body_text)
            entries_layout.addWidget(entry_frame)
            entries_added += 1

        # Store checkboxes reference
        setattr(self, f"popup_{key}_{section_key}_imported_checkboxes", checkboxes)

        return entries_added > 0

    def _update_h4_subsection_from_checkboxes(self, section_key: str):
        """
        Update the H4 subsection text field based on checked imported data checkboxes.
        """
        key = 'h4'
        checkboxes = getattr(self, f"popup_{key}_{section_key}_imported_checkboxes", [])
        text_widget = getattr(self, f"popup_{key}_{section_key}_text_widget", None)

        if not text_widget:
            return

        # Collect text from checked boxes
        checked_texts = []
        for cb in checkboxes:
            if cb.isChecked():
                full_text = cb.property("full_text")
                if full_text:
                    # Get just the first 500 chars as a summary
                    summary = full_text[:500].strip()
                    if len(full_text) > 500:
                        summary += "..."
                    checked_texts.append(summary)

        # Update the text widget with checked content
        if checked_texts:
            text_widget.setPlainText("\n\n---\n\n".join(checked_texts))
        else:
            text_widget.clear()
