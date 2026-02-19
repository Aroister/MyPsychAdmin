# ================================================================
#  TRIBUNAL REPORT PAGE — Psychiatric Tribunal Report Writer
# ================================================================

from __future__ import annotations

import re
import sys
from datetime import datetime, timedelta
from html import unescape

from PySide6.QtCore import Qt, Signal, QSize, QEvent
from PySide6.QtGui import QColor, QFontDatabase
from PySide6.QtWidgets import QGraphicsDropShadowEffect
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QSplitter, QStackedWidget, QTextEdit,
    QSizePolicy, QPushButton, QToolButton, QComboBox, QColorDialog
)
from mypsy_richtext_editor import MyPsychAdminRichTextEditor
from utils.resource_path import resource_path
from spell_check_textedit import enable_spell_check_on_textedit


# ================================================================
# ZOOM HELPER FUNCTION
# ================================================================

def create_zoom_row(text_edit: QTextEdit, base_size: int = 12) -> QHBoxLayout:
    """Create a zoom controls row for any QTextEdit."""
    zoom_row = QHBoxLayout()
    zoom_row.setSpacing(2)
    zoom_row.addStretch()

    text_edit._font_size = base_size

    zoom_out_btn = QPushButton("−")
    zoom_out_btn.setFixedSize(16, 16)
    zoom_out_btn.setCursor(Qt.CursorShape.PointingHandCursor)
    zoom_out_btn.setToolTip("Decrease font size")
    zoom_out_btn.setStyleSheet("""
        QPushButton {
            background: #f3f4f6;
            color: #374151;
            border: 1px solid #d1d5db;
            border-radius: 3px;
            font-size: 11px;
            font-weight: bold;
        }
        QPushButton:hover { background: #e5e7eb; }
    """)
    zoom_row.addWidget(zoom_out_btn)

    zoom_in_btn = QPushButton("+")
    zoom_in_btn.setFixedSize(16, 16)
    zoom_in_btn.setCursor(Qt.CursorShape.PointingHandCursor)
    zoom_in_btn.setToolTip("Increase font size")
    zoom_in_btn.setStyleSheet("""
        QPushButton {
            background: #f3f4f6;
            color: #374151;
            border: 1px solid #d1d5db;
            border-radius: 3px;
            font-size: 11px;
            font-weight: bold;
        }
        QPushButton:hover { background: #e5e7eb; }
    """)
    zoom_row.addWidget(zoom_in_btn)

    # Get the current stylesheet pattern
    current_style = text_edit.styleSheet()

    def zoom_in():
        text_edit._font_size = min(text_edit._font_size + 2, 28)
        new_style = re.sub(r'font-size:\s*\d+px', f'font-size: {text_edit._font_size}px', current_style)
        text_edit.setStyleSheet(new_style)

    def zoom_out():
        text_edit._font_size = max(text_edit._font_size - 2, 8)
        new_style = re.sub(r'font-size:\s*\d+px', f'font-size: {text_edit._font_size}px', current_style)
        text_edit.setStyleSheet(new_style)

    zoom_in_btn.clicked.connect(zoom_in)
    zoom_out_btn.clicked.connect(zoom_out)

    return zoom_row


# ================================================================
# TRIBUNAL TOOLBAR
# ================================================================

class TribunalToolbar(QWidget):
    """Toolbar for the Tribunal Report Page (without magnifying glass, organise, upload)."""

    # Formatting signals
    set_font_family = Signal(str)
    set_font_size = Signal(int)

    toggle_bold = Signal()
    toggle_italic = Signal()
    toggle_underline = Signal()

    set_text_color = Signal(QColor)
    set_highlight_color = Signal(QColor)

    set_align_left = Signal()
    set_align_center = Signal()
    set_align_right = Signal()
    set_align_justify = Signal()

    bullet_list = Signal()
    numbered_list = Signal()

    indent = Signal()
    outdent = Signal()

    undo = Signal()
    redo = Signal()

    insert_date = Signal()
    insert_section_break = Signal()

    export_docx = Signal()
    check_spelling = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)

        self.setFixedHeight(64)  # Reduced by 20%
        self.setStyleSheet("""
            TribunalToolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 15px;
                font-weight: 500;
            }
            QToolButton:hover {
                background: rgba(0,0,0,0.08);
            }
            QComboBox {
                background: rgba(255,255,255,0.85);
                color: #333333;
                border: 1px solid rgba(0,0,0,0.15);
                border-radius: 6px;
                padding: 4px 8px;
                font-size: 14px;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox QAbstractItemView {
                background: white;
                color: #333333;
                selection-background-color: #e0e0e0;
            }
        """)

        # Outer layout for the toolbar
        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        # Scroll area to prevent compression
        scroll = QScrollArea()
        scroll.setWidgetResizable(False)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOn)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(64)  # Reduced by 20%

        # Container widget for toolbar content
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(60)  # Reduced by 20%
        container.setMinimumWidth(1200)  # Force scrollbar when viewport is smaller
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 8, 2)
        layout.setSpacing(10)

        # ---------------------------------------------------------
        # EXPORT BUTTON - PROMINENT STYLING
        # ---------------------------------------------------------
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(150, 42)
        export_btn.setStyleSheet("""
            QToolButton {
                background: #2563eb;
                color: white;
                font-size: 18px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #1d4ed8;
            }
            QToolButton:pressed {
                background: #1e40af;
            }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addWidget(export_btn)

        # ---------------------------------------------------------
        # UPLOADED DOCS BUTTON (dropdown menu)
        # ---------------------------------------------------------
        from PySide6.QtWidgets import QMenu
        import_btn = QToolButton()
        import_btn.setText("Uploaded Docs")
        import_btn.setFixedSize(170, 42)
        import_btn.setPopupMode(QToolButton.InstantPopup)
        self.upload_menu = QMenu()
        import_btn.setMenu(self.upload_menu)
        import_btn.setStyleSheet("""
            QToolButton {
                background: #10b981;
                color: white;
                font-size: 18px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 8px 16px;
            }
            QToolButton:hover {
                background: #059669;
            }
            QToolButton:pressed {
                background: #047857;
            }
            QToolButton::menu-indicator { image: none; }
        """)
        layout.addWidget(import_btn)

        # ---------------------------------------------------------
        # FONT FAMILY
        # ---------------------------------------------------------
        self.font_combo = QComboBox()
        self.font_combo.setFixedWidth(180)

        families = QFontDatabase.families()
        if sys.platform == "win32":
            preferred = ["Segoe UI", "Calibri", "Cambria", "Arial", "Times New Roman"]
        else:
            preferred = ["Avenir Next", "Avenir", "SF Pro Text", "Helvetica Neue", "Helvetica"]

        added = set()
        for f in preferred:
            if f in families:
                self.font_combo.addItem(f)
                added.add(f)
        for f in families:
            if f not in added:
                self.font_combo.addItem(f)

        self.font_combo.currentTextChanged.connect(self.set_font_family.emit)
        layout.addWidget(self.font_combo)

        # ---------------------------------------------------------
        # FONT SIZE
        # ---------------------------------------------------------
        self.size_combo = QComboBox()
        self.size_combo.setFixedWidth(65)
        for sz in [8, 9, 10, 11, 12, 14, 16, 18, 20, 22]:
            self.size_combo.addItem(str(sz))
        self.size_combo.currentTextChanged.connect(lambda v: self.set_font_size.emit(int(v)))
        layout.addWidget(self.size_combo)

        # Simple button helper
        def btn(label, slot):
            b = QToolButton()
            b.setText(label)
            b.setMinimumWidth(36)
            b.clicked.connect(slot)
            return b

        # ---------------------------------------------------------
        # BASIC STYLES
        # ---------------------------------------------------------
        layout.addWidget(btn("B", self.toggle_bold.emit))
        layout.addWidget(btn("I", self.toggle_italic.emit))
        layout.addWidget(btn("U", self.toggle_underline.emit))

        # ---------------------------------------------------------
        # COLORS
        # ---------------------------------------------------------
        layout.addWidget(btn("A", self._choose_text_color))
        layout.addWidget(btn("🖍", self._choose_highlight_color))

        # ---------------------------------------------------------
        # ALIGNMENT
        # ---------------------------------------------------------
        layout.addWidget(btn("L", self.set_align_left.emit))
        layout.addWidget(btn("C", self.set_align_center.emit))
        layout.addWidget(btn("R", self.set_align_right.emit))
        layout.addWidget(btn("J", self.set_align_justify.emit))

        # ---------------------------------------------------------
        # LISTS / INDENTATION
        # ---------------------------------------------------------
        layout.addWidget(btn("•", self.bullet_list.emit))
        layout.addWidget(btn("1.", self.numbered_list.emit))
        layout.addWidget(btn("→", self.indent.emit))
        layout.addWidget(btn("←", self.outdent.emit))

        # ---------------------------------------------------------
        # UNDO / REDO
        # ---------------------------------------------------------
        layout.addWidget(btn("⟲", self.undo.emit))
        layout.addWidget(btn("⟳", self.redo.emit))

        # ---------------------------------------------------------
        # INSERTS
        # ---------------------------------------------------------
        layout.addWidget(btn("Date", self.insert_date.emit))
        layout.addWidget(btn("Break", self.insert_section_break.emit))

        # ---------------------------------------------------------
        # SPELL CHECK
        # ---------------------------------------------------------
        spell_btn = QToolButton()
        spell_btn.setText("Spell Check")
        spell_btn.setFixedSize(120, 38)
        spell_btn.setStyleSheet("""
            QToolButton {
                background: #f59e0b;
                color: white;
                font-size: 15px;
                font-weight: 600;
                border: none;
                border-radius: 8px;
                padding: 6px 12px;
            }
            QToolButton:hover { background: #d97706; }
            QToolButton:pressed { background: #b45309; }
        """)
        spell_btn.setToolTip("Jump to next spelling error")
        spell_btn.clicked.connect(self.check_spelling.emit)
        layout.addWidget(spell_btn)

        # Finalize scroll area
        scroll.setWidget(container)
        outer_layout.addWidget(scroll)

    def _choose_text_color(self):
        col = QColorDialog.getColor(QColor("black"), self)
        if col.isValid():
            self.set_text_color.emit(col)

    def _choose_highlight_color(self):
        col = QColorDialog.getColor(QColor("yellow"), self)
        if col.isValid():
            self.set_highlight_color.emit(col)


# ================================================================
# CARD WIDGET
# ================================================================

class TribunalCardWidget(QFrame):
    """A clickable card for a tribunal report section."""

    clicked = Signal(str)  # Emits the section key

    STYLE_NORMAL = """
        TribunalCardWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 12px;
            padding: 16px;
        }
        TribunalCardWidget:hover {
            border-color: #8b5cf6;
            background: #faf5ff;
        }
    """

    STYLE_SELECTED = """
        TribunalCardWidget {
            background: #ede9fe;
            border: 2px solid #8b5cf6;
            border-left: 4px solid #7c3aed;
            border-radius: 12px;
            padding: 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(8)

        # Title
        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("""
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        layout.addWidget(title_lbl)

        # Editor (rich text with formatting support)
        self.editor = MyPsychAdminRichTextEditor()
        self.editor.setPlaceholderText("Click to edit...")
        self.editor.setReadOnly(False)
        self._editor_height = 80  # Reduced by 20%
        self.editor.setMinimumHeight(48)  # Reduced by 20%
        self.editor.setMaximumHeight(self._editor_height)
        self.editor.setStyleSheet("""
            QTextEdit {
                background: #f9fafb;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 8px;
                font-size: 16px;
                color: #374151;
            }
        """)
        editor_zoom = create_zoom_row(self.editor, base_size=16)
        layout.addLayout(editor_zoom)
        layout.addWidget(self.editor)

        # Expand/resize bar
        self.expand_bar = QFrame()
        self.expand_bar.setFixedHeight(10)  # Reduced by 20%
        self.expand_bar.setCursor(Qt.CursorShape.SizeVerCursor)
        self.expand_bar.setStyleSheet("""
            QFrame {
                background: #e5e7eb;
                border-radius: 2px;
                margin: 4px 40px;
            }
            QFrame:hover {
                background: #8b5cf6;
            }
        """)
        self.expand_bar.installEventFilter(self)
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0
        layout.addWidget(self.expand_bar)

    def eventFilter(self, obj, event):
        """Handle drag events on the expand bar."""
        if obj == self.expand_bar:
            if event.type() == QEvent.Type.MouseButtonPress:
                self._dragging = True
                self._drag_start_y = event.globalPosition().y()
                self._drag_start_height = self._editor_height
                return True
            elif event.type() == QEvent.Type.MouseMove and self._dragging:
                delta = event.globalPosition().y() - self._drag_start_y
                new_height = max(60, min(500, self._drag_start_height + delta))
                self._editor_height = int(new_height)
                self.editor.setMinimumHeight(self._editor_height)
                self.editor.setMaximumHeight(self._editor_height)
                self.editor.setFixedHeight(self._editor_height)
                return True
            elif event.type() == QEvent.Type.MouseButtonRelease:
                self._dragging = False
                return True
        return super().eventFilter(obj, event)

    def mousePressEvent(self, event):
        # Only emit if not clicking on the editor or expand bar
        if not self.editor.geometry().contains(event.pos()) and not self.expand_bar.geometry().contains(event.pos()):
            self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        """Set the selected state of the card."""
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        """Return whether the card is selected."""
        return self._selected


# ================================================================
# TRIBUNAL HEADING WIDGET (for Yes/No sections with no card content)
# ================================================================

class TribunalHeadingWidget(QFrame):
    """A clickable heading for tribunal sections that don't need an editor card."""

    clicked = Signal(str)  # Emits the section key

    STYLE_NORMAL = """
        TribunalHeadingWidget {
            background: white;
            border: 1px solid #e0e0e0;
            border-radius: 8px;
            padding: 8px 16px;
        }
        TribunalHeadingWidget:hover {
            border-color: #8b5cf6;
            background: #faf5ff;
        }
    """

    STYLE_SELECTED = """
        TribunalHeadingWidget {
            background: #ede9fe;
            border: 2px solid #8b5cf6;
            border-left: 4px solid #7c3aed;
            border-radius: 8px;
            padding: 8px 16px;
        }
    """

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.title = title
        self._selected = False

        self.setFrameShape(QFrame.Shape.StyledPanel)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet(self.STYLE_NORMAL)
        # Allow height to adjust based on content, but expand horizontally
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum)
        self.setMinimumHeight(40)  # Reduced by 20%

        # Add subtle shadow
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(12)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 25))
        self.setGraphicsEffect(shadow)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)

        # Title - with word wrap enabled
        self.title_lbl = QLabel(title)
        self.title_lbl.setWordWrap(True)
        self.title_lbl.setStyleSheet("""
            font-size: 17px;
            font-weight: 600;
            color: #1f2937;
            background: transparent;
        """)
        self.title_lbl.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        layout.addWidget(self.title_lbl, 1)

        # Click indicator arrow (fixed width, aligned to right)
        arrow = QLabel("\u25B6")  # Right arrow
        arrow.setStyleSheet("""
            font-size: 12px;
            color: #8b5cf6;
            background: transparent;
        """)
        arrow.setFixedWidth(20)
        arrow.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        layout.addWidget(arrow)

        # Dummy editor attribute for compatibility with card-based code
        self.editor = _DummyEditor()

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def setSelected(self, selected: bool):
        """Set the selected state."""
        if self._selected != selected:
            self._selected = selected
            self.setStyleSheet(self.STYLE_SELECTED if selected else self.STYLE_NORMAL)

    def isSelected(self) -> bool:
        """Return whether the heading is selected."""
        return self._selected


class _DummyEditor:
    """Minimal dummy editor for TribunalHeadingWidget compatibility."""
    def toPlainText(self):
        return ""
    def setPlainText(self, text):
        pass
    def toHtml(self):
        return ""
    def setHtml(self, html):
        pass
    def clear(self):
        pass


# ================================================================
# RECENT ADMISSION POPUP (for Section 8)
# ================================================================

class RecentAdmissionPopup(QWidget):
    """Popup showing filtered entries around the current admission date."""

    sent = Signal(str)  # Emits the text content

    def __init__(self, parent=None):
        super().__init__(parent)
        self._entries = []
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # Title
        title = QLabel("Recent Admission Data")
        title.setStyleSheet("""
            font-size: 16px;
            font-weight: 700;
            color: #7c3aed;
        """)
        layout.addWidget(title)

        # Subtitle with date range info
        self.date_info = QLabel("Entries around current admission")
        self.date_info.setStyleSheet("""
            font-size: 12px;
            color: #6b7280;
            font-style: italic;
        """)
        layout.addWidget(self.date_info)

        # Scrollable content area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                background: white;
            }
        """)

        # Content widget
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(12, 12, 12, 12)
        self.content_layout.setSpacing(8)
        scroll.setWidget(self.content_widget)

        layout.addWidget(scroll, 1)

        # Summary frame (collapsible, hidden by default)
        self.summary_frame = QFrame()
        self.summary_frame.setStyleSheet("""
            QFrame {
                background: #fef3c7;
                border: 1px solid #f59e0b;
                border-radius: 8px;
                padding: 8px;
            }
        """)
        self.summary_frame.hide()
        summary_layout = QVBoxLayout(self.summary_frame)
        summary_layout.setContentsMargins(8, 8, 8, 8)

        summary_header = QHBoxLayout()
        summary_label = QLabel("Summary")
        summary_label.setStyleSheet("font-weight: 600; color: #92400e;")
        summary_header.addWidget(summary_label)
        summary_header.addStretch()

        # Copy button
        copy_summary_btn = QPushButton("Copy")
        copy_summary_btn.setFixedWidth(50)
        copy_summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 4px 8px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background: #d97706;
            }
        """)
        copy_summary_btn.clicked.connect(self._copy_summary)
        summary_header.addWidget(copy_summary_btn)

        # Close X button
        close_btn = QPushButton("✕")
        close_btn.setFixedWidth(24)
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #92400e;
                border: none;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                color: #78350f;
                background: rgba(0,0,0,0.1);
                border-radius: 4px;
            }
        """)
        close_btn.clicked.connect(lambda: self.summary_frame.hide())
        summary_header.addWidget(close_btn)

        summary_layout.addLayout(summary_header)

        self.summary_text = QTextEdit()
        self.summary_text.setMaximumHeight(120)  # Reduced by 20%
        self.summary_text.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #fcd34d;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        summary_zoom = create_zoom_row(self.summary_text, base_size=12)
        summary_layout.addLayout(summary_zoom)
        summary_layout.addWidget(self.summary_text)
        layout.addWidget(self.summary_frame)

        # Button row
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        # Clear button
        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #f3f4f6;
                color: #374151;
                border: 1px solid #d1d5db;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #e5e7eb;
            }
        """)
        clear_btn.clicked.connect(self._clear)
        btn_layout.addWidget(clear_btn)

        # Summary button
        self.summary_btn = QPushButton("Summary")
        self.summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #d97706;
            }
        """)
        self.summary_btn.clicked.connect(self._generate_summary)
        btn_layout.addWidget(self.summary_btn)

        # Send to report button
        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #8b5cf6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #7c3aed;
            }
        """)
        send_btn.clicked.connect(self._send_to_letter)
        btn_layout.addWidget(send_btn)

        layout.addLayout(btn_layout)

    def set_entries(self, entries: list, date_range_info: str = ""):
        """Set the entries to display."""
        self._entries = entries

        # Update date info label
        if date_range_info:
            self.date_info.setText(date_range_info)

        # Clear existing content
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not entries:
            no_data = QLabel("No entries found for this time period")
            no_data.setStyleSheet("color: #9ca3af; font-style: italic;")
            self.content_layout.addWidget(no_data)
            return

        # Add each entry
        for entry in entries:
            entry_frame = QFrame()
            entry_frame.setStyleSheet("""
                QFrame {
                    background: #f9fafb;
                    border: 1px solid #e5e7eb;
                    border-radius: 6px;
                    padding: 8px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(8, 8, 8, 8)
            entry_layout.setSpacing(4)

            # Date header if available
            date_str = entry.get("date", "")
            if date_str:
                date_label = QLabel(str(date_str))
                date_label.setStyleSheet("""
                    font-size: 11px;
                    font-weight: 600;
                    color: #7c3aed;
                """)
                entry_layout.addWidget(date_label)

            # Entry text
            text = entry.get("text", "")
            text_label = QLabel(text)
            text_label.setWordWrap(True)
            text_label.setStyleSheet("""
                font-size: 12px;
                color: #374151;
            """)
            entry_layout.addWidget(text_label)

            self.content_layout.addWidget(entry_frame)

        self.content_layout.addStretch()

    def _clear(self):
        """Clear all entries."""
        self._entries = []
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        self.summary_frame.hide()

    def _generate_summary(self):
        """Generate a concise narrative summary of admission circumstances."""
        if not self._entries:
            self.summary_text.setPlainText("No entries to summarize.")
            self.summary_frame.show()
            return

        import re

        # Combine all text for searching
        all_text = ""
        for entry in self._entries:
            text = entry.get("text", "") or entry.get("content", "")
            if text:
                all_text += text + "\n\n"

        summary_parts = []

        # Step 1: Look for "Reason for Admission" section (keep only the longest/most detailed)
        reason_matches = re.findall(r'[Rr]eason\s+for\s+[Aa]dmission[:\s]*([^.]+(?:\.[^.]+){0,5})', all_text)
        if reason_matches:
            # Get the longest one (most detailed)
            best_reason = max(reason_matches, key=len)
            reason = " ".join(best_reason.split())
            summary_parts.append(f"Reason for Admission: {reason}")

        # Step 2: Extract key events by category - keep only the BEST (longest) of each type
        EVENT_CATEGORIES = {
            'arrested': [
                r'[Aa]rrested\s+(?:on|and|for|by)[^.]{10,200}\.',
            ],
            'pos': [
                r'[Pp]lace\s+of\s+[Ss]afety[^.]{5,150}\.',
                r'POS\s+[^.]{5,100}\.',
            ],
            's136': [
                r'[Ss]ection\s*136[^.]{5,150}\.',
                r'[Ss]136[^.]{5,100}\.',
            ],
            'mhaa': [
                r'[Rr]eferred\s+for\s+(?:a\s+)?MHAA[^.]{5,150}\.',
            ],
            'detained': [
                r'[Dd]etained\s+under\s+[Ss](?:ection\s*)?\d[^.]{5,100}\.',
            ],
        }

        # For each category, find all matches and keep only the longest/most detailed one
        for category, patterns in EVENT_CATEGORIES.items():
            all_matches = []
            for pattern in patterns:
                matches = re.findall(pattern, all_text, re.IGNORECASE)
                all_matches.extend(matches)

            if all_matches:
                # Keep only the longest match (most detailed)
                best_match = max(all_matches, key=len)
                clean = " ".join(best_match.split())

                # Don't add if it's essentially already in summary
                is_duplicate = False
                for existing in summary_parts:
                    # Check if key phrases overlap
                    existing_lower = existing.lower()
                    clean_lower = clean.lower()
                    # If >50% of words overlap, consider it duplicate
                    existing_words = set(existing_lower.split())
                    clean_words = set(clean_lower.split())
                    if len(existing_words & clean_words) > len(clean_words) * 0.5:
                        is_duplicate = True
                        break

                if not is_duplicate:
                    summary_parts.append(clean)

        # Build final summary (limit to 6 items for readability)
        if summary_parts:
            summary = "\n".join(summary_parts[:6])
        else:
            summary = "No admission-related details found in the filtered notes."

        self.summary_text.setPlainText(summary)
        self.summary_frame.show()

    def _copy_summary(self):
        """Copy summary to clipboard."""
        from PySide6.QtWidgets import QApplication
        clipboard = QApplication.clipboard()
        clipboard.setText(self.summary_text.toPlainText())

    def _send_to_letter(self):
        """Send content to the letter card."""
        # Combine all entry texts
        texts = []
        for entry in self._entries:
            text = entry.get("text", "").strip()
            date_str = entry.get("date", "")
            if text:
                if date_str:
                    texts.append(f"[{date_str}] {text}")
                else:
                    texts.append(text)

        combined = "\n\n".join(texts)
        self.sent.emit(combined)


# ================================================================
# FIXED DATA PANEL - Generic panel for displaying extracted data
# ================================================================

class FixedDataPanel(QWidget):
    """Generic fixed panel for displaying extracted data with summary."""

    sent = Signal(str)  # Emits the text content

    def __init__(self, title: str, subtitle: str = "", parent=None):
        super().__init__(parent)
        self._entries = []
        self._title = title
        self._subtitle = subtitle
        self.notes = []  # Store raw notes for other sections to access
        self._setup_ui()

    def _setup_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        # Title
        title = QLabel(self._title)
        title.setStyleSheet("""
            font-size: 16px;
            font-weight: 700;
            color: #7c3aed;
        """)
        layout.addWidget(title)

        # Subtitle
        self.subtitle_label = QLabel(self._subtitle or "Extracted data from notes")
        self.subtitle_label.setStyleSheet("""
            font-size: 12px;
            color: #6b7280;
            font-style: italic;
        """)
        layout.addWidget(self.subtitle_label)

        # Scrollable content area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid #e5e7eb;
                border-radius: 8px;
                background: white;
            }
        """)

        # Content widget
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(12, 12, 12, 12)
        self.content_layout.setSpacing(8)
        scroll.setWidget(self.content_widget)

        layout.addWidget(scroll, 1)

        # Summary frame (collapsible, hidden by default)
        self.summary_frame = QFrame()
        self.summary_frame.setStyleSheet("""
            QFrame {
                background: #fef3c7;
                border: 1px solid #f59e0b;
                border-radius: 8px;
                padding: 8px;
            }
        """)
        self.summary_frame.hide()
        summary_layout = QVBoxLayout(self.summary_frame)
        summary_layout.setContentsMargins(8, 8, 8, 8)

        summary_header = QHBoxLayout()
        summary_label = QLabel("Summary")
        summary_label.setStyleSheet("font-weight: 600; color: #92400e;")
        summary_header.addWidget(summary_label)
        summary_header.addStretch()

        # Copy button
        copy_summary_btn = QPushButton("Copy")
        copy_summary_btn.setFixedWidth(50)
        copy_summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 4px 8px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background: #d97706;
            }
        """)
        copy_summary_btn.clicked.connect(self._copy_summary)
        summary_header.addWidget(copy_summary_btn)

        # Close X button
        close_btn = QPushButton("✕")
        close_btn.setFixedWidth(24)
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #92400e;
                border: none;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                color: #78350f;
                background: rgba(0,0,0,0.1);
                border-radius: 4px;
            }
        """)
        close_btn.clicked.connect(lambda: self.summary_frame.hide())
        summary_header.addWidget(close_btn)

        summary_layout.addLayout(summary_header)

        self.summary_text = QTextEdit()
        self.summary_text.setMaximumHeight(120)  # Reduced by 20%
        self.summary_text.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #fcd34d;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        summary_zoom = create_zoom_row(self.summary_text, base_size=12)
        summary_layout.addLayout(summary_zoom)
        summary_layout.addWidget(self.summary_text)
        layout.addWidget(self.summary_frame)

        # Button row
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(12)

        # Clear button
        clear_btn = QPushButton("Clear")
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #f3f4f6;
                color: #374151;
                border: 1px solid #d1d5db;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #e5e7eb;
            }
        """)
        clear_btn.clicked.connect(self._clear)
        btn_layout.addWidget(clear_btn)

        # Summary button
        self.summary_btn = QPushButton("Summary")
        self.summary_btn.setStyleSheet("""
            QPushButton {
                background: #f59e0b;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: #d97706;
            }
        """)
        self.summary_btn.clicked.connect(self._generate_summary)
        btn_layout.addWidget(self.summary_btn)

        # Send to report button
        send_btn = QPushButton("Send to Report")
        send_btn.setStyleSheet("""
            QPushButton {
                background: #8b5cf6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 8px;
                font-weight: 600;
            }
            QPushButton:hover {
                background: #7c3aed;
            }
        """)
        send_btn.clicked.connect(self._send_to_letter)
        btn_layout.addWidget(send_btn)

        layout.addLayout(btn_layout)

    def set_entries(self, entries: list, info_text: str = ""):
        """Set the entries to display."""
        self._entries = entries
        self.notes = entries  # Store for other sections to access

        # Update subtitle
        if info_text:
            self.subtitle_label.setText(info_text)

        # Clear existing content
        while self.content_layout.count():
            item = self.content_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        if not entries:
            no_data = QLabel("No entries found. Use Import File to load data.")
            no_data.setStyleSheet("color: #9ca3af; font-style: italic;")
            self.content_layout.addWidget(no_data)
            return

        # For incident panels (sections 17 & 18), show summary directly instead of all entries
        if "harm" in self._title.lower() or "property" in self._title.lower():
            self._show_incident_summary_directly()
            return

        # Limit UI display to first 100 entries for performance (all entries still available for summary)
        display_entries = entries[:100]
        if len(entries) > 100:
            # Add note about truncation
            note = QLabel(f"Showing first 100 of {len(entries)} entries. Click 'Generate Summary' to see all.")
            note.setStyleSheet("color: #6366f1; font-style: italic; font-weight: 600; padding: 8px;")
            self.content_layout.addWidget(note)

        # Add each entry
        for entry in display_entries:
            entry_frame = QFrame()
            entry_frame.setStyleSheet("""
                QFrame {
                    background: #f9fafb;
                    border: 1px solid #e5e7eb;
                    border-radius: 6px;
                    padding: 8px;
                }
            """)
            entry_layout = QVBoxLayout(entry_frame)
            entry_layout.setContentsMargins(8, 8, 8, 8)
            entry_layout.setSpacing(4)

            # Date header if available
            date_str = entry.get("date", "") or entry.get("datetime", "")
            if date_str:
                date_label = QLabel(str(date_str))
                date_label.setStyleSheet("""
                    font-size: 11px;
                    font-weight: 600;
                    color: #6b7280;
                """)
                entry_layout.addWidget(date_label)

            # Content text
            text = entry.get("text", "") or entry.get("content", "")
            if text:
                # Truncate for display
                display_text = text[:500] + "..." if len(text) > 500 else text
                text_label = QLabel(display_text)
                text_label.setWordWrap(True)
                text_label.setStyleSheet("font-size: 12px; color: #374151;")
                entry_layout.addWidget(text_label)

            self.content_layout.addWidget(entry_frame)

        # Add stretch at end
        self.content_layout.addStretch()

    def _clear(self):
        """Clear the entries."""
        self.set_entries([])
        self.summary_frame.hide()

    def _generate_summary(self):
        """Generate a smart summary based on panel type."""
        if not self._entries:
            self.summary_text.setPlainText("No data to summarize.")
            self.summary_frame.show()
            return

        import re
        from datetime import datetime

        # Combine all text for analysis
        all_text = ""
        for entry in self._entries:
            text = entry.get("text", "") or entry.get("content", "")
            if text:
                all_text += text + "\n\n"

        if not all_text.strip():
            self.summary_text.setPlainText("No text content to summarize.")
            self.summary_frame.show()
            return

        summary_parts = []

        # Smart extraction based on panel title
        if "Previous" in self._title or "Admission" in self._title or "Psychiatric" in self._title:
            # =================================================================
            # PSYCHIATRIC HISTORY SUMMARY - Extract full narrative paragraphs
            # =================================================================

            # Helper to normalize date for deduplication (returns month/year string)
            def normalize_date(text):
                """Extract month/year from date text for deduplication."""
                # Try DD/MM/YY format
                m = re.search(r'(\d{1,2})/(\d{1,2})/(\d{2,4})', text)
                if m:
                    day, month, year = m.groups()
                    if len(year) == 2:
                        year = '20' + year if int(year) < 50 else '19' + year
                    return f"{month}/{year}"
                # Try "26 Dec 2011" format
                m = re.search(r'(\d{1,2})\s+(\w{3})\w*\s+(\d{4})', text)
                if m:
                    day, month_name, year = m.groups()
                    month_map = {'jan': '1', 'feb': '2', 'mar': '3', 'apr': '4', 'may': '5', 'jun': '6',
                                 'jul': '7', 'aug': '8', 'sep': '9', 'oct': '10', 'nov': '11', 'dec': '12'}
                    month = month_map.get(month_name.lower()[:3], '0')
                    return f"{month}/{year}"
                # Try "December 2011" format
                m = re.search(r'(\w+)\s+(\d{4})', text)
                if m:
                    month_name, year = m.groups()
                    month_map = {'january': '1', 'february': '2', 'march': '3', 'april': '4',
                                 'may': '5', 'june': '6', 'july': '7', 'august': '8',
                                 'september': '9', 'october': '10', 'november': '11', 'december': '12'}
                    month = month_map.get(month_name.lower(), '0')
                    if month != '0':
                        return f"{month}/{year}"
                return None

            def parse_date_for_sort(text):
                """Parse date for chronological sorting."""
                # Try DD/MM/YY format
                m = re.search(r'(\d{1,2})/(\d{1,2})/(\d{2,4})', text)
                if m:
                    day, month, year = m.groups()
                    if len(year) == 2:
                        year = '20' + year if int(year) < 50 else '19' + year
                    try:
                        return datetime(int(year), int(month), int(day))
                    except:
                        pass
                # Try "26 Dec 2011" format
                m = re.search(r'(\d{1,2})\s+(\w{3})\w*\s+(\d{4})', text)
                if m:
                    day, month_name, year = m.groups()
                    month_map = {'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
                                 'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12}
                    month = month_map.get(month_name.lower()[:3], 1)
                    try:
                        return datetime(int(year), month, int(day))
                    except:
                        pass
                # Try "December 2011" format
                m = re.search(r'(\w+)\s+(\d{4})', text)
                if m:
                    month_name, year = m.groups()
                    month_map = {'january': 1, 'february': 2, 'march': 3, 'april': 4,
                                 'may': 5, 'june': 6, 'july': 7, 'august': 8,
                                 'september': 9, 'october': 10, 'november': 11, 'december': 12}
                    month = month_map.get(month_name.lower(), 0)
                    if month:
                        try:
                            return datetime(int(year), month, 1)
                        except:
                            pass
                return datetime.min

            # STEP 1: Find the FIRST detailed "Past Psych Hx" entry with full paragraphs
            narrative_paragraphs = []
            covered_dates = set()  # Track dates covered by narrative

            # Look for entries with "Past Psych Hx" or detailed admission narratives
            for entry in self._entries:
                text = entry.get("text", "") or entry.get("content", "")
                if not text:
                    continue

                # Split into lines and look for dated narrative lines (>80 chars = narrative, not table)
                lines = text.split('\n')
                dated_narrative_lines = []

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    # Check if line starts with a date pattern and has substantial text
                    # Date patterns: "December 2011 (4 day)", "26/12/2011:", "5/2/2012:"
                    is_dated = bool(re.match(
                        r'^(?:\d{1,2}/\d{1,2}/\d{2,4}|'  # DD/MM/YY or DD/MM/YYYY
                        r'\w+\s+\d{4}\s*\([^)]+\)|'       # December 2011 (4 day)
                        r'X\d+\s+in-patient)',            # X3 in-patient
                        line
                    ))

                    # Narrative lines are >80 chars (table rows are shorter)
                    if is_dated and len(line) > 80:
                        dated_narrative_lines.append(line)
                        # Track this date as covered
                        date_key = normalize_date(line)
                        if date_key:
                            covered_dates.add(date_key)

                    # Also extract "Previous" lines
                    if line.lower().startswith('previous '):
                        dated_narrative_lines.append(line)

                # If we found multiple narrative lines, this is our detailed entry
                if len(dated_narrative_lines) >= 2:
                    narrative_paragraphs = dated_narrative_lines
                    break  # Use only the first detailed entry

            # Sort narrative paragraphs chronologically
            narrative_paragraphs.sort(key=parse_date_for_sort)

            # Add narrative paragraphs to summary
            for para in narrative_paragraphs:
                summary_parts.append(para)

            # STEP 2: Extract admission history table (structured table format)
            admission_table_rows = []
            admission_block = re.search(
                r'[Aa]dmission\s+[Hh]istory:?\s*((?:\d{1,2}\s+\w{3}\s+\d{4}[^\n]*\n?)+)',
                all_text
            )
            if admission_block:
                lines = admission_block.group(1).strip().split('\n')
                for line in lines:
                    line = line.strip()
                    if line and re.match(r'\d', line):
                        # Check if this date is already covered by narrative
                        date_key = normalize_date(line)
                        if date_key and date_key not in covered_dates:
                            admission_table_rows.append(line)

            # Sort table rows chronologically
            admission_table_rows.sort(key=parse_date_for_sort)

            # Add table rows that aren't duplicates
            for row in admission_table_rows:
                summary_parts.append(row)

            # STEP 3: Find current status from most recent entry
            current_status_lines = []
            status_patterns = [
                r'[Hh]e\s+is\s+(?:well\s+known|currently)[^.]+\.',
                r'[Ss]he\s+is\s+(?:well\s+known|currently)[^.]+\.',
                r'[Hh]e\s+has\s+been\s+discharged[^.]+(?:CTO|community)[^.]*\.',
                r'[Ss]he\s+has\s+been\s+discharged[^.]+(?:CTO|community)[^.]*\.',
                r'[Hh]e\s+is\s+currently\s+under[^.]+\.',
                r'[Ss]he\s+is\s+currently\s+under[^.]+\.',
            ]

            # Check most recent entries for current status
            for entry in reversed(self._entries):
                text = entry.get("text", "") or entry.get("content", "")
                for pattern in status_patterns:
                    matches = re.findall(pattern, text)
                    for match in matches:
                        clean = match.strip()
                        if clean and clean not in current_status_lines:
                            current_status_lines.append(clean)
                if current_status_lines:
                    break

            # Add current status at the end
            for status in current_status_lines[:3]:
                if status not in summary_parts:
                    summary_parts.append(status)

        elif "Progress" in self._title or "Mental State" in self._title:
            # Use risk-scored summary if available
            summary_parts = self._generate_risk_based_summary()
            if summary_parts:
                # Build final summary and return early
                summary = "\n\n".join(summary_parts)
                self.summary_text.setPlainText(summary)
                self.summary_frame.show()
                return

        elif "Incident" in self._title or "harm" in self._title.lower() or "property" in self._title.lower():
            # Entries are already pre-processed - group by date and join
            summary_parts.append("INCIDENTS")
            summary_parts.append("-" * 40)

            if self._entries:
                import re
                from datetime import datetime

                # Group all entries by date
                date_to_entries = {}
                for entry in self._entries:
                    date_key = entry.get('date', 'Unknown')
                    text = entry.get('text', '') or entry.get('content', '')
                    content = text.split(':', 1)[1].strip() if ':' in text else text

                    if date_key not in date_to_entries:
                        date_to_entries[date_key] = []
                    if content and content not in date_to_entries[date_key]:
                        date_to_entries[date_key].append(content)

                # Sort by date (most recent first)
                def parse_date_key(date_str):
                    try:
                        return datetime.strptime(date_str, '%d/%m/%Y')
                    except:
                        return datetime.min

                sorted_dates = sorted(date_to_entries.keys(), key=parse_date_key, reverse=True)

                # Filter and join entries per date
                incident_terms = ['aggressive', 'aggression', 'violent', 'violence', 'threaten',
                                  'abusive', 'agitation', 'restrain', 'seclusion', 'secluded',
                                  'fight', 'fist', 'threw', 'throw', 'scream', 'intox', 'arrest',
                                  'confrontational', 'barricade', 'custody', 'retaliation', 'alarming',
                                  'inappropriat', 'self-neglect', 'sexual', 'stripped']

                filtered_count = 0
                seen_content = set()  # Track seen content for de-duplication

                for date_key in sorted_dates:
                    filtered_entries = []
                    for content in date_to_entries[date_key]:
                        content_lower = content.lower()

                        if content_lower.startswith('diagnosis:') or content_lower.startswith('positive behaviour'):
                            filtered_count += 1
                            continue
                        if content_lower.startswith('to self:') or content_lower.startswith('to others:'):
                            filtered_count += 1
                            continue
                        if content_lower.startswith('risk:') or content_lower.startswith('risks'):
                            filtered_count += 1
                            continue
                        if content_lower.startswith('self neglect:') or content_lower.startswith('self-neglect:'):
                            filtered_count += 1
                            continue
                        if 'without incident' in content_lower or 'no evidence' in content_lower or 'nothing to indicate' in content_lower:
                            filtered_count += 1
                            continue
                        if 'risk of' in content_lower or 'risk to' in content_lower:
                            filtered_count += 1
                            continue
                        if 'medication for agitation' in content_lower:
                            filtered_count += 1
                            continue
                        if 'call police if' in content_lower:
                            filtered_count += 1
                            continue
                        if 'less agitation' in content_lower or 'less aggression' in content_lower:
                            filtered_count += 1
                            continue
                        if 'reduced agitation' in content_lower or 'reduced aggression' in content_lower:
                            filtered_count += 1
                            continue
                        if 'police and ambulance to be called if' in content_lower or 'police to be called if' in content_lower:
                            filtered_count += 1
                            continue
                        if '(agitation)' in content_lower or '(agitation' in content_lower:
                            filtered_count += 1
                            continue
                        if 'previous' in content_lower:
                            filtered_count += 1
                            continue
                        if 'threatened to walk out' in content_lower:
                            filtered_count += 1
                            continue
                        if 'can be aggressive' in content_lower:
                            filtered_count += 1
                            continue
                        if 'police to be called if' in content_lower:
                            filtered_count += 1
                            continue
                        if re.search(r'\bnil\b', content_lower) or re.search(r'\bnon\b', content_lower):
                            filtered_count += 1
                            continue
                        if re.search(r'\bno\b', content_lower):
                            filtered_count += 1
                            continue
                        if re.search(r'\b(did not|didn\'t|does not)\b(?:\s+\S+){0,5}\s+aggressive', content_lower):
                            filtered_count += 1
                            continue
                        if re.search(r'aggressive(?:\s+\S+){0,5}\s+\b(did not|didn\'t|does not)\b', content_lower):
                            filtered_count += 1
                            continue

                        # Skip if "not" within 10 words of any incident term
                        skip_not_proximity = False
                        if re.search(r'\bnot\b', content_lower):
                            words = content_lower.split()
                            not_positions = [i for i, w in enumerate(words) if w == 'not']
                            for not_pos in not_positions:
                                for term in incident_terms:
                                    for i, word in enumerate(words):
                                        if term in word and abs(i - not_pos) <= 10:
                                            skip_not_proximity = True
                                            break
                                    if skip_not_proximity:
                                        break
                                if skip_not_proximity:
                                    break
                        if skip_not_proximity:
                            filtered_count += 1
                            continue

                        # De-duplication: skip if content already seen
                        content_normalized = ' '.join(content_lower.split())
                        if content_normalized in seen_content:
                            filtered_count += 1
                            continue
                        seen_content.add(content_normalized)

                        filtered_entries.append(content)

                    if filtered_entries:
                        joined_text = ' | '.join(filtered_entries)
                        summary_parts.append(f"  • {date_key}: {joined_text}")

                print(f"[INCIDENT SUMMARY] Grouped to {len(sorted_dates)} dates, filtered {filtered_count} items")
            else:
                summary_parts.append("  Nil recorded.")

            summary = "\n".join(summary_parts)
            self.summary_text.setPlainText(summary)
            self.summary_frame.show()
            return

        # Fallback: extract first meaningful sentences
        if not summary_parts:
            for entry in self._entries[:10]:
                text = entry.get("text", "") or entry.get("content", "")
                if text:
                    sentences = text.split('.')
                    for s in sentences[:2]:
                        s = s.strip()
                        if len(s) > 30 and len(summary_parts) < 5:
                            summary_parts.append(s[:200])

        # Build final summary
        summary = "\n".join(summary_parts)

        self.summary_text.setPlainText(summary if summary else "No meaningful content found to summarize.")
        self.summary_frame.show()

    def _show_incident_summary_directly(self):
        """Show incident summary directly in the content area instead of individual entries."""
        import re
        from datetime import datetime

        # Group all entries by date - join all entries for same date
        date_to_entries = {}
        for entry in self._entries:
            date_key = entry.get('date', 'Unknown')
            text = entry.get('text', '') or entry.get('content', '')
            # Remove the date prefix from text if present (format: "DD/MM/YYYY: content")
            if ':' in text:
                content = text.split(':', 1)[1].strip()
            else:
                content = text

            if date_key not in date_to_entries:
                date_to_entries[date_key] = []
            if content and content not in date_to_entries[date_key]:  # Avoid duplicates
                date_to_entries[date_key].append(content)

        # Sort by date (most recent first)
        def parse_date_key(date_str):
            try:
                return datetime.strptime(date_str, '%d/%m/%Y')
            except:
                return datetime.min

        sorted_dates = sorted(date_to_entries.keys(), key=parse_date_key, reverse=True)

        # Filter entries and collect summary lines (one per date with all entries joined)
        summary_lines = []
        filtered_count = 0
        seen_content = set()  # Track seen content for de-duplication across dates

        for date_key in sorted_dates:
            entries_for_date = date_to_entries[date_key]
            filtered_entries = []

            # Incident terms for proximity check
            incident_terms = ['aggressive', 'aggression', 'violent', 'violence', 'threaten',
                              'abusive', 'agitation', 'restrain', 'seclusion', 'secluded',
                              'fight', 'fist', 'threw', 'throw', 'scream', 'intox', 'arrest',
                              'confrontational', 'barricade', 'custody', 'retaliation', 'alarming',
                              'inappropriat', 'self-neglect', 'sexual', 'stripped']

            for content in entries_for_date:
                content_lower = content.lower()

                # Skip lines starting with "Diagnosis:", "Positive Behaviour", "To self:", "To others:", "Risk:", "Risks", "Self neglect:"
                if content_lower.startswith('diagnosis:') or content_lower.startswith('positive behaviour'):
                    filtered_count += 1
                    continue
                if content_lower.startswith('to self:') or content_lower.startswith('to others:'):
                    filtered_count += 1
                    continue
                if content_lower.startswith('risk:') or content_lower.startswith('risks'):
                    filtered_count += 1
                    continue
                if content_lower.startswith('self neglect:') or content_lower.startswith('self-neglect:'):
                    filtered_count += 1
                    continue

                # Skip lines containing these phrases
                if 'without incident' in content_lower or 'no evidence' in content_lower or 'nothing to indicate' in content_lower:
                    filtered_count += 1
                    continue

                if 'risk of' in content_lower or 'risk to' in content_lower:
                    filtered_count += 1
                    continue

                if 'medication for agitation' in content_lower:
                    filtered_count += 1
                    continue

                if 'call police if' in content_lower:
                    filtered_count += 1
                    continue

                # Skip lines with less/reduced agitation/aggression
                if 'less agitation' in content_lower or 'less aggression' in content_lower:
                    filtered_count += 1
                    continue
                if 'reduced agitation' in content_lower or 'reduced aggression' in content_lower:
                    filtered_count += 1
                    continue

                # Skip police-related conditionals
                if 'police and ambulance to be called if' in content_lower or 'police to be called if' in content_lower:
                    filtered_count += 1
                    continue

                # Skip if agitation in brackets
                if '(agitation)' in content_lower or '(agitation' in content_lower:
                    filtered_count += 1
                    continue

                # Skip lines with 'previous'
                if 'previous' in content_lower:
                    filtered_count += 1
                    continue

                # Skip 'threatened to walk out'
                if 'threatened to walk out' in content_lower:
                    filtered_count += 1
                    continue

                # Skip 'can be aggressive'
                if 'can be aggressive' in content_lower:
                    filtered_count += 1
                    continue

                # Skip 'Police to be called if' (case variations)
                if 'police to be called if' in content_lower:
                    filtered_count += 1
                    continue

                # Skip lines containing "nil" or "non" as whole words
                if re.search(r'\bnil\b', content_lower) or re.search(r'\bnon\b', content_lower):
                    filtered_count += 1
                    continue

                # Skip lines containing the word "no" as a whole word
                if re.search(r'\bno\b', content_lower):
                    filtered_count += 1
                    continue

                # Skip if "did not", "didn't", "does not" within 5 words of "aggressive"
                if re.search(r'\b(did not|didn\'t|does not)\b(?:\s+\S+){0,5}\s+aggressive', content_lower):
                    filtered_count += 1
                    continue
                if re.search(r'aggressive(?:\s+\S+){0,5}\s+\b(did not|didn\'t|does not)\b', content_lower):
                    filtered_count += 1
                    continue

                # Skip if "not" within 10 words of any incident term
                skip_not_proximity = False
                if re.search(r'\bnot\b', content_lower):
                    words = content_lower.split()
                    not_positions = [i for i, w in enumerate(words) if w == 'not']
                    for not_pos in not_positions:
                        for term in incident_terms:
                            for i, word in enumerate(words):
                                if term in word and abs(i - not_pos) <= 10:
                                    skip_not_proximity = True
                                    break
                            if skip_not_proximity:
                                break
                        if skip_not_proximity:
                            break
                if skip_not_proximity:
                    filtered_count += 1
                    continue

                # De-duplication: skip if content already seen (normalize for comparison)
                content_normalized = ' '.join(content_lower.split())  # Normalize whitespace
                if content_normalized in seen_content:
                    filtered_count += 1
                    continue
                seen_content.add(content_normalized)

                filtered_entries.append(content)

            # Join all filtered entries for this date
            if filtered_entries:
                joined_text = ' | '.join(filtered_entries)
                summary_lines.append(f"{date_key}: {joined_text}")

        print(f"[INCIDENT DIRECT] Showing {len(summary_lines)} date entries (filtered {filtered_count} individual items)")

        # Update subtitle with correct filtered count
        self.subtitle_label.setText(f"{len(summary_lines)} dates with incidents")

        # Display summary lines directly in content area
        if summary_lines:
            # Add each incident line
            for line in summary_lines:
                line_label = QLabel(f"• {line}")
                line_label.setWordWrap(True)
                line_label.setStyleSheet("""
                    font-size: 12px;
                    color: #374151;
                    padding: 4px 8px;
                    background: #f9fafb;
                    border-left: 3px solid #6366f1;
                    margin: 2px 0;
                """)
                self.content_layout.addWidget(line_label)
        else:
            no_data = QLabel("No relevant incidents found after filtering.")
            no_data.setStyleSheet("color: #9ca3af; font-style: italic; padding: 8px;")
            self.content_layout.addWidget(no_data)

        self.content_layout.addStretch()

        # Hide the summary button and frame since we're showing summary directly
        self.summary_btn.hide()
        self.summary_frame.hide()

    def _generate_risk_based_summary(self):
        """Generate narrative summary with Progress, Capacity, and Insight sections."""
        import re
        from datetime import datetime
        from pathlib import Path

        if not self._entries:
            return []

        # Load risk dictionary for scoring (used internally, not displayed)
        risk_dict = {}
        risk_file = Path(__file__).parent / "riskDICT.txt"
        if risk_file.exists():
            with open(risk_file, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line or ',' not in line:
                        continue
                    parts = line.rsplit(',', 1)
                    term = parts[0].strip().lower()
                    score_str = parts[1].strip() if len(parts) > 1 else ''
                    if term and score_str:
                        try:
                            risk_dict[term] = int(score_str)
                        except:
                            pass

        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        def get_highest_scoring_line(content, risk_dict):
            """Return the line with the highest risk score (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            best_line = None
            best_score = -1
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                line_score = sum(points for term, points in risk_dict.items() if term in line_lower)
                if line_score > best_score:
                    best_score = line_score
                    best_line = cleaned
            return best_line

        def get_relevant_keyword_line(content, keywords):
            """Extract the most relevant line containing keywords (must be >= 100 chars)."""
            lines = content.strip().split('\n')
            for line in lines:
                cleaned = line.strip()
                if len(cleaned) < 100:
                    continue
                line_lower = cleaned.lower()
                for kw in keywords:
                    if kw in line_lower:
                        return cleaned
            return None

        # Categorize entries
        entries_data = []
        for entry in self._entries:
            content = entry.get('content', '') or entry.get('text', '') or ''
            content_lower = content.lower()
            if not content:
                continue

            note_date = parse_date(entry.get('date') or entry.get('datetime'))
            date_str = note_date.strftime('%d/%m/%Y') if note_date else 'Unknown'

            # Calculate internal score
            score = sum(points for term, points in risk_dict.items() if term in content_lower)

            entries_data.append({
                'date': note_date,
                'date_str': date_str,
                'score': score,
                'content': content,
                'content_lower': content_lower
            })

        # Sort by date
        entries_data.sort(key=lambda x: x['date'] or datetime.min)

        # === EXTRACT EVENTS BY CATEGORY ===

        # High concern events (threshold 1500)
        concern_events = []
        for e in entries_data:
            if e['score'] >= 1500:
                concern_events.append(e)

        # Violence/Aggression events
        violence_keywords = ['violence', 'violent', 'assault', 'attack', 'fight', 'aggression', 'aggressive', 'physical altercation']
        violence_events = []
        for e in entries_data:
            if any(kw in e['content_lower'] for kw in violence_keywords):
                violence_events.append(e)

        # DNA / Lack of contact
        dna_keywords = ['dna', 'did not attend', 'no contact', 'uncontactable', 'failed to attend', 'missed appointment', 'not seen', 'no response']
        dna_events = []
        for e in entries_data:
            if any(kw in e['content_lower'] for kw in dna_keywords):
                dna_events.append(e)

        # Positive progress
        positive_keywords = ['stable', 'settled', 'calm', 'pleasant', 'appropriate', 'well presented', 'good rapport', 'engaging well', 'cooperative', 'progress']
        positive_events = []
        for e in entries_data:
            if any(kw in e['content_lower'] for kw in positive_keywords) and e['score'] < 100:
                positive_events.append(e)

        # Capacity assessments
        capacity_keywords = ['capacity', 'capacitous', 'lacks capacity', 'has capacity', 'mental capacity', 'mca', 'decision-making']
        capacity_events = []
        for e in entries_data:
            if any(kw in e['content_lower'] for kw in capacity_keywords):
                capacity_events.append(e)

        # Insight records
        insight_keywords = ['insight', 'awareness', 'understands', 'accepts diagnosis', 'denies illness', 'lacks insight', 'good insight', 'poor insight', 'partial insight']
        insight_events = []
        for e in entries_data:
            if any(kw in e['content_lower'] for kw in insight_keywords):
                insight_events.append(e)

        # === ENGAGEMENT / ACTIVITIES EVENTS ===
        # General engagement keywords
        engagement_keywords = ['engage', 'engaged', 'engagement', 'did not engage', 'no engagement', 'refused to engage', 'engaged well', 'good engagement', 'poor engagement', 'minimal engagement']
        engagement_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in engagement_keywords)]

        # OT, Groups, Section 17 leave
        ot_keywords = ['occupational therapy', 'ot session', 'ot group', 'o.t.', 'occupational therapist']
        ot_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in ot_keywords)]

        groups_keywords = ['group', 'groups', 'attended group', 'psychology group', 'therapy group', 'art group', 'music group', 'community meeting', 'ward round']
        groups_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in groups_keywords)]

        s17_keywords = ['section 17', 's17', 'leave', 'escorted leave', 'unescorted leave', 'ground leave', 'community leave', 'home leave', 'overnight leave']
        s17_events = [e for e in entries_data if any(kw in e['content_lower'] for kw in s17_keywords)]

        # === BUILD NARRATIVE SUMMARY ===
        summary_parts = []

        # Date range header
        dates = [e['date'] for e in entries_data if e['date']]
        if dates:
            earliest = min(dates).strftime('%d/%m/%Y')
            latest = max(dates).strftime('%d/%m/%Y')

        # === SECTION 1: PROGRESS ===
        summary_parts.append("PROGRESS")
        summary_parts.append("-" * 40)

        if dates:
            summary_parts.append(f"Review period: {earliest} to {latest}")

        def get_first_line(content):
            """Get first non-empty line as note type identifier."""
            for line in content.strip().split('\n'):
                cleaned = line.strip()
                if cleaned:
                    return cleaned[:60] + "..." if len(cleaned) > 60 else cleaned
            return "Note"

        def extract_relevant_text(content, keywords):
            """Extract sentence/phrase containing keyword (must be >= 100 chars)."""
            content_lower = content.lower()
            for kw in keywords:
                pos = content_lower.find(kw)
                if pos != -1:
                    # Find sentence boundaries
                    start = max(0, content.rfind('.', 0, pos) + 1)
                    end = content.find('.', pos)
                    if end == -1:
                        end = min(len(content), pos + 150)
                    else:
                        end = min(end + 1, pos + 200)
                    excerpt = content[start:end].strip()
                    if len(excerpt) >= 100:
                        return excerpt
            return ""

        # Concern events - show all with highest scoring line
        if concern_events:
            sorted_concern = sorted(concern_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"Significant behavioural concerns noted ({len(sorted_concern)}):")
            for e in sorted_concern:
                relevant_line = get_highest_scoring_line(e['content'], risk_dict)
                if relevant_line:
                    summary_parts.append(f"  • {e['date_str']}: {relevant_line}")
        else:
            summary_parts.append("No significant behavioural concerns in this period.")

        # Violence/Aggression (threshold 2000) - show all with relevant line
        summary_parts.append("")
        summary_parts.append("Episodes of violence or aggression:")
        high_violence = [e for e in violence_events if e['score'] >= 2000]
        if high_violence:
            sorted_violence = sorted(high_violence, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_violence)} episodes)")
            for e in sorted_violence:
                relevant_line = get_highest_scoring_line(e['content'], risk_dict)
                if relevant_line:
                    summary_parts.append(f"  • {e['date_str']}: {relevant_line}")
        else:
            summary_parts.append("  Nil noted.")

        # DNAs - show all with relevant line
        if dna_events:
            sorted_dna = sorted(dna_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append("")
            summary_parts.append(f"Missed appointments or lack of contact ({len(sorted_dna)}):")
            for e in sorted_dna:
                relevant_line = get_relevant_keyword_line(e['content'], dna_keywords)
                if relevant_line:
                    summary_parts.append(f"  • {e['date_str']}: {relevant_line}")

        # Positive progress - summarize to key descriptors only
        def summarize_positive(content_lower):
            """Summarize positive entry to key descriptor."""
            if 'superficially settled' in content_lower or 'superfically settled' in content_lower:
                return 'settled (superficially)'
            if 'settled' in content_lower or 'calm' in content_lower:
                return 'settled'
            if 'stable' in content_lower:
                return 'stable'
            if 'pleasant' in content_lower:
                return 'pleasant'
            if 'cooperative' in content_lower:
                return 'cooperative'
            if 'engaging well' in content_lower or 'good rapport' in content_lower:
                return 'engaging well'
            if 'well presented' in content_lower or 'appropriate' in content_lower:
                return 'well presented'
            if 'progress' in content_lower:
                return 'progress noted'
            return 'positive'

        if positive_events:
            sorted_positive = sorted(positive_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append("")
            summary_parts.append(f"Positive observations ({len(sorted_positive)}):")
            for e in sorted_positive:
                descriptor = summarize_positive(e['content_lower'])
                summary_parts.append(f"  • {e['date_str']}: {descriptor}")

        # === SECTION 2: CAPACITY ===
        summary_parts.append("")
        summary_parts.append("CAPACITY")
        summary_parts.append("-" * 40)
        if capacity_events:
            sorted_capacity = sorted(capacity_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_capacity)} records)")
            for e in sorted_capacity:
                note_type = get_first_line(e['content'])
                capacity_text = extract_relevant_text(e['content'], capacity_keywords)
                summary_parts.append(f"  • Capacity mentioned on {e['date_str']} ({note_type})")
                if capacity_text:
                    summary_parts.append(f"    {capacity_text}")
        else:
            summary_parts.append("  Not recorded.")

        # === SECTION 3: INSIGHT ===
        summary_parts.append("")
        summary_parts.append("INSIGHT")
        summary_parts.append("-" * 40)
        if insight_events:
            sorted_insight = sorted(insight_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_insight)} records)")
            for e in sorted_insight:
                note_type = get_first_line(e['content'])
                insight_text = extract_relevant_text(e['content'], insight_keywords)
                summary_parts.append(f"  • Insight mentioned on {e['date_str']} ({note_type})")
                if insight_text:
                    summary_parts.append(f"    {insight_text}")
        else:
            summary_parts.append("  Not recorded.")

        # === SECTION 4: ENGAGEMENT / ACTIVITIES ===
        summary_parts.append("")
        summary_parts.append("")
        summary_parts.append("ENGAGEMENT / ACTIVITIES")
        summary_parts.append("-" * 40)

        # Summarization function for engagement
        def summarize_engagement(content_lower):
            """Summarize engagement entry to key descriptor."""
            if 'did not engage' in content_lower or 'no engagement' in content_lower or 'nil engagement' in content_lower:
                return 'no engagement'
            if 'refused to engage' in content_lower or 'declined to engage' in content_lower:
                return 'refused to engage'
            if 'minimal engagement' in content_lower or 'limited engagement' in content_lower:
                return 'minimal engagement'
            if 'engaged well' in content_lower or 'good engagement' in content_lower:
                return 'engaged well'
            if 'engaged on needs' in content_lower or 'needs basis' in content_lower or 'needs led' in content_lower or 'needy bas' in content_lower:
                return 'engaged on needs basis'
            if 'polite when engag' in content_lower or 'pleasant when engag' in content_lower:
                return 'polite when engaged'
            if 'lack of engagement' in content_lower or 'unable to assess' in content_lower:
                return 'lack of engagement'
            if 'engaged' in content_lower:
                return 'engaged'
            return 'engagement noted'

        # General Engagement
        summary_parts.append("")
        summary_parts.append("General Engagement:")
        if engagement_events:
            sorted_engagement = sorted(engagement_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_engagement)} records)")
            for e in sorted_engagement:
                descriptor = summarize_engagement(e['content_lower'])
                summary_parts.append(f"  • {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No engagement records found.")

        # Summarization function for OT
        def summarize_ot(content_lower):
            """Summarize OT entry to key descriptor."""
            if 'refused' in content_lower or 'declined' in content_lower:
                return 'declined OT'
            if 'engaged well' in content_lower or 'participated' in content_lower:
                return 'engaged in OT'
            if 'ot session' in content_lower or 'occupational therapy' in content_lower:
                return 'OT session'
            return 'OT noted'

        # OT
        summary_parts.append("")
        summary_parts.append("Occupational Therapy:")
        if ot_events:
            sorted_ot = sorted(ot_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_ot)} records)")
            for e in sorted_ot:
                descriptor = summarize_ot(e['content_lower'])
                summary_parts.append(f"  • {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No OT sessions recorded.")

        # Summarization function for groups
        def summarize_groups(content_lower):
            """Summarize groups entry to key descriptor."""
            if 'did not attend' in content_lower or 'refused' in content_lower or 'declined' in content_lower:
                return 'did not attend'
            if 'ward round' in content_lower:
                return 'ward round'
            if 'community meeting' in content_lower:
                return 'community meeting'
            if 'psychology group' in content_lower or 'therapy group' in content_lower:
                return 'therapy group'
            if 'art group' in content_lower:
                return 'art group'
            if 'music group' in content_lower:
                return 'music group'
            if 'attended group' in content_lower or 'group session' in content_lower:
                return 'attended group'
            if 'group room' in content_lower:
                return 'group room session'
            return 'group activity'

        # Groups
        summary_parts.append("")
        summary_parts.append("Groups / Ward Activities:")
        if groups_events:
            sorted_groups = sorted(groups_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_groups)} records)")
            for e in sorted_groups:
                descriptor = summarize_groups(e['content_lower'])
                summary_parts.append(f"  • {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No group activities recorded.")

        # Summarization function for S17 leave
        def summarize_leave(content_lower):
            """Summarize S17 leave entry to key descriptor."""
            # Check for NON-USE of leave FIRST (most specific patterns)
            if ('leave: none' in content_lower or 'leave: nil' in content_lower or
                'leave: not utilised' in content_lower or 'leave: not used' in content_lower or
                'leave: non used' in content_lower or 'leave not utilised' in content_lower or
                'leave not used' in content_lower or 'did not use leave' in content_lower or
                'no leave today' in content_lower or 'leave not taken' in content_lower or
                'did not utilise leave' in content_lower or 'none utilised' in content_lower or
                'did not access leave' in content_lower or 'no leave' in content_lower):
                return 'leave not used'
            if 'refused leave' in content_lower or 'declined leave' in content_lower:
                return 'refused leave'
            if 'leave could not' in content_lower or 'unable to facilitate' in content_lower:
                return 'leave not facilitated'
            # Check for ACTUAL USE - must have clear indicators of leave being taken
            # Shopping leave - but check it was actually used
            if ('shopping leave' in content_lower or 'went shopping' in content_lower or
                'to the shop' in content_lower or 'marston green' in content_lower or
                'chelmsley wood' in content_lower) and 'leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'shopping leave'
            # Community leave - check it was used
            if 'community leave' in content_lower:
                if 'went on community' in content_lower or 'utilised community' in content_lower or 'accessed community' in content_lower:
                    return 'community leave'
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'community leave'
            # Ground leave - check it was used
            if 'ground leave' in content_lower or 'grounds leave' in content_lower:
                if ('went on' in content_lower or 'utilised' in content_lower or 'accessed' in content_lower or
                    'facilitated' in content_lower or 'took' in content_lower or 'had' in content_lower):
                    return 'ground leave'
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'ground leave'
            # Escorted leave - check it was used
            if 'escorted leave' in content_lower:
                if ('went on' in content_lower or 'utilised' in content_lower or 'accessed' in content_lower or
                    'facilitated' in content_lower or 'took' in content_lower):
                    return 'escorted leave'
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'escorted leave'
            # Unescorted leave
            if 'unescorted leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'unescorted leave'
            # Home leave
            if 'home leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'home leave'
            # Overnight leave
            if 'overnight leave' in content_lower:
                if 'refused' not in content_lower and 'declined' not in content_lower and 'not' not in content_lower:
                    return 'overnight leave'
            # Clear indicators that leave was actually taken
            if ('went on leave' in content_lower or 'utilised leave' in content_lower or
                'accessed leave' in content_lower or 'leave was facilitated' in content_lower or
                'leave went well' in content_lower or 'took him on leave' in content_lower or
                'took her on leave' in content_lower or 'escorted on leave' in content_lower or
                'used his leave' in content_lower or 'used her leave' in content_lower or
                'utilised his leave' in content_lower or 'utilised her leave' in content_lower or
                'leave to' in content_lower and ('shop' in content_lower or 'walk' in content_lower)):
                return 'leave utilised'
            # Default: if just "leave" mentioned without clear use indicator, assume not used
            return 'leave not used'

        # Section 17 Leave
        summary_parts.append("")
        summary_parts.append("Section 17 Leave:")
        if s17_events:
            sorted_s17 = sorted(s17_events, key=lambda x: x['date'] or datetime.min, reverse=True)
            summary_parts.append(f"  ({len(sorted_s17)} records)")
            for e in sorted_s17:
                descriptor = summarize_leave(e['content_lower'])
                summary_parts.append(f"  • {e['date_str']}: {descriptor}")
        else:
            summary_parts.append("  No Section 17 leave recorded.")

        return summary_parts

    def _copy_summary(self):
        """Copy summary to clipboard."""
        from PySide6.QtWidgets import QApplication
        clipboard = QApplication.clipboard()
        clipboard.setText(self.summary_text.toPlainText())

    def _send_to_letter(self):
        """Send content to the letter card."""
        import re
        from datetime import datetime

        # For incident panels (17 & 18), use grouped format matching the display
        if "harm" in self._title.lower() or "property" in self._title.lower():
            # Group all entries by date
            date_to_entries = {}
            for entry in self._entries:
                date_key = entry.get('date', 'Unknown')
                text = entry.get('text', '') or entry.get('content', '')
                content = text.split(':', 1)[1].strip() if ':' in text else text

                if date_key not in date_to_entries:
                    date_to_entries[date_key] = []
                if content and content not in date_to_entries[date_key]:
                    date_to_entries[date_key].append(content)

            # Sort by date (most recent first)
            def parse_date_key(date_str):
                try:
                    return datetime.strptime(date_str, '%d/%m/%Y')
                except:
                    return datetime.min

            sorted_dates = sorted(date_to_entries.keys(), key=parse_date_key, reverse=True)

            # Filter and join entries per date
            incident_terms = ['aggressive', 'aggression', 'violent', 'violence', 'threaten',
                              'abusive', 'agitation', 'restrain', 'seclusion', 'secluded',
                              'fight', 'fist', 'threw', 'throw', 'scream', 'intox', 'arrest',
                              'confrontational', 'barricade', 'custody', 'retaliation', 'alarming',
                              'inappropriat', 'self-neglect', 'sexual', 'stripped']

            texts = []
            seen_content = set()  # Track seen content for de-duplication

            for date_key in sorted_dates:
                filtered_entries = []
                for content in date_to_entries[date_key]:
                    content_lower = content.lower()

                    if content_lower.startswith('diagnosis:') or content_lower.startswith('positive behaviour'):
                        continue
                    if content_lower.startswith('to self:') or content_lower.startswith('to others:'):
                        continue
                    if content_lower.startswith('risk:') or content_lower.startswith('risks'):
                        continue
                    if content_lower.startswith('self neglect:') or content_lower.startswith('self-neglect:'):
                        continue
                    if 'without incident' in content_lower or 'no evidence' in content_lower or 'nothing to indicate' in content_lower:
                        continue
                    if 'risk of' in content_lower or 'risk to' in content_lower:
                        continue
                    if 'medication for agitation' in content_lower:
                        continue
                    if 'call police if' in content_lower:
                        continue
                    if 'less agitation' in content_lower or 'less aggression' in content_lower:
                        continue
                    if 'reduced agitation' in content_lower or 'reduced aggression' in content_lower:
                        continue
                    if 'police and ambulance to be called if' in content_lower or 'police to be called if' in content_lower:
                        continue
                    if '(agitation)' in content_lower or '(agitation' in content_lower:
                        continue
                    if 'previous' in content_lower:
                        continue
                    if 'threatened to walk out' in content_lower:
                        continue
                    if 'can be aggressive' in content_lower:
                        continue
                    if 'police to be called if' in content_lower:
                        continue
                    if re.search(r'\bnil\b', content_lower) or re.search(r'\bnon\b', content_lower):
                        continue
                    if re.search(r'\bno\b', content_lower):
                        continue
                    if re.search(r'\b(did not|didn\'t|does not)\b(?:\s+\S+){0,5}\s+aggressive', content_lower):
                        continue
                    if re.search(r'aggressive(?:\s+\S+){0,5}\s+\b(did not|didn\'t|does not)\b', content_lower):
                        continue

                    # Skip if "not" within 10 words of any incident term
                    skip_not_proximity = False
                    if re.search(r'\bnot\b', content_lower):
                        words = content_lower.split()
                        not_positions = [i for i, w in enumerate(words) if w == 'not']
                        for not_pos in not_positions:
                            for term in incident_terms:
                                for i, word in enumerate(words):
                                    if term in word and abs(i - not_pos) <= 10:
                                        skip_not_proximity = True
                                        break
                                if skip_not_proximity:
                                    break
                            if skip_not_proximity:
                                break
                    if skip_not_proximity:
                        continue

                    # De-duplication: skip if content already seen
                    content_normalized = ' '.join(content_lower.split())
                    if content_normalized in seen_content:
                        continue
                    seen_content.add(content_normalized)

                    filtered_entries.append(content)

                if filtered_entries:
                    joined_text = ' | '.join(filtered_entries)
                    texts.append(f"{date_key}: {joined_text}")

            combined = "\n\n".join(texts)
            self.sent.emit(combined)
            return

        # Default behavior for other panels
        texts = []
        for entry in self._entries:
            text = (entry.get("text", "") or entry.get("content", "")).strip()
            date_str = entry.get("date", "") or entry.get("datetime", "")
            if text:
                if date_str:
                    texts.append(f"[{date_str}] {text}")
                else:
                    texts.append(text)

        combined = "\n\n".join(texts)
        self.sent.emit(combined)


# ================================================================
# TRIBUNAL REPORT PAGE
# ================================================================

class TribunalReportPage(QWidget):
    """Main page for creating Psychiatric Tribunal Reports."""

    go_back = Signal()  # Signal to return to reports page

    # Sections based on T131 form - exact questions with numbers
    SECTIONS = [
        ("1. Patient Details", "patient_details"),
        ("2. Name of Responsible Clinician", "author"),
        ("3. Factors affecting understanding or ability to cope with hearing", "factors_hearing"),
        ("4. Adjustments for tribunal to consider", "adjustments"),
        ("5. Index offence(s) and relevant forensic history", "forensic"),
        ("6. Dates of previous involvement with mental health services", "previous_mh_dates"),
        ("7. Reasons for previous admission or recall to hospital", "previous_admission_reasons"),
        ("8. Circumstances leading to current admission", "current_admission"),
        ("9. Mental disorder and diagnosis", "diagnosis"),
        ("10. Learning disability", "learning_disability"),
        ("11. Mental disorder requiring detention", "detention_required"),
        ("12. Medical treatment prescribed, provided, offered or planned", "treatment"),
        ("13. Strengths or positive factors", "strengths"),
        ("14. Current progress, behaviour, capacity and insight", "progress"),
        ("15. What is the patient's understanding of, compliance with", "compliance"),
        ("16. Deprivation of liberty under MCA 2005 consideration", "mca_dol"),
        ("17. Incidents of harm to self or others", "risk_harm"),
        ("18. Incidents of property damage", "risk_property"),
        ("19. Section 2: Detention justified for health, safety or protection", "s2_detention"),
        ("20. Other sections: Medical treatment justified for health, safety or protection", "other_detention"),
        ("21. Risk if discharged from hospital", "discharge_risk"),
        ("22. Community risk management", "community"),
        ("23. Recommendations to tribunal", "recommendations"),
        ("24. Signature", "signature"),
    ]

    def __init__(self, parent=None, db=None):
        super().__init__(parent)
        self.db = db
        self.cards = {}
        self.popups = {}
        self.popup_memory = {}
        self._active_editor = None
        self._current_gender = "Male"  # Default gender
        self._selected_card_key = None  # Track currently selected card
        self._my_details = self._load_my_details()

        # Store extracted data at page level for popups created later
        self._extracted_raw_notes = []
        self._extracted_categories = {}

        # Bidirectional sync flag to prevent recursive updates
        self._syncing = False

        # Guard flags to prevent reprocessing on navigation
        self._data_processed_id = None
        self._notes_processed_id = None

        self._setup_ui()

        # Connect to shared store for cross-talk with nursing form
        self._connect_shared_store()

    def _load_my_details(self) -> dict:
        """Load clinician details from database."""
        if not self.db:
            return {}

        details = self.db.get_clinician_details()
        if not details:
            return {}

        # details is a tuple: (id, full_name, role_title, discipline, registration_body,
        #                      registration_number, phone, email, team_service,
        #                      hospital_org, ward_department, signature_block)
        return {
            "full_name": details[1] or "",
            "role_title": details[2] or "",
            "discipline": details[3] or "",
            "registration_body": details[4] or "",
            "registration_number": details[5] or "",
            "phone": details[6] or "",
            "email": details[7] or "",
            "team_service": details[8] or "",
            "hospital_org": details[9] or "",
            "ward_department": details[10] or "",
            "signature_block": details[11] or "",
        }

    def _connect_shared_store(self):
        """Connect to SharedDataStore for cross-report data sharing."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            shared_store.report_sections_changed.connect(self._on_report_sections_changed)
            shared_store.notes_changed.connect(self._on_notes_changed)
            shared_store.extracted_data_changed.connect(self._on_extracted_data_changed)
            shared_store.patient_info_changed.connect(self._on_patient_info_changed)
            print("[TRIBUNAL] Connected to SharedDataStore signals (sections, notes, extracted_data, patient_info)")

            # Check if there's already data in the store (uploaded before this form existed)
            self._check_shared_store_for_existing_data()
        except Exception as e:
            print(f"[TRIBUNAL] Failed to connect to SharedDataStore: {e}")

    def _check_shared_store_for_existing_data(self):
        """Check SharedDataStore for existing data when page is created."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()

            # Check for existing report sections (cross-talk)
            existing_sections = shared_store.report_sections
            source = shared_store.get_report_source()
            if existing_sections and source != "tribunal":
                print(f"[TRIBUNAL] Found existing sections from {source}, populating...")
                self._on_report_sections_changed(existing_sections, source)

            # Check for existing patient info
            patient_info = shared_store.patient_info
            if patient_info and any(patient_info.values()):
                print(f"[TRIBUNAL] Found existing patient info in SharedDataStore")
                self._fill_patient_details(patient_info)

            # Check for existing notes (skip if report data present)
            notes = shared_store.notes
            if notes and not self._has_report_data():
                print(f"[TRIBUNAL] Found {len(notes)} existing notes in SharedDataStore")
                self._extracted_raw_notes = notes

            # Check for existing extracted data (goes through guarded handler)
            extracted_data = shared_store.extracted_data
            if extracted_data:
                print(f"[TRIBUNAL] Found existing extracted data in SharedDataStore")
                self._on_extracted_data_changed(extracted_data)
        except Exception as e:
            print(f"[TRIBUNAL] Error checking shared store: {e}")

    def _on_patient_info_changed(self, patient_info: dict):
        """Handle patient info updates from SharedDataStore."""
        if patient_info and any(patient_info.values()):
            print(f"[TRIBUNAL] Received patient info from SharedDataStore: {list(k for k,v in patient_info.items() if v)}")
            self._fill_patient_details(patient_info)

    def _has_report_data(self):
        """Check if report data has been imported (local or via SharedDataStore)."""
        if hasattr(self, '_imported_report_data') and self._imported_report_data:
            return True
        # Also check SharedDataStore for report sections from another form
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            source = shared_store.get_report_source()
            if source and source != "tribunal" and shared_store.report_sections:
                return True
        except Exception:
            pass
        return False

    def _on_notes_changed(self, notes: list):
        """Handle notes updates from SharedDataStore."""
        if notes:
            # Skip if report data exists (report takes priority)
            if self._has_report_data():
                print(f"[TRIBUNAL] Skipping notes from SharedDataStore - report data already imported")
                return
            print(f"[TRIBUNAL] Received {len(notes)} notes from SharedDataStore")
            self._extracted_raw_notes = notes
            # Refresh popups that use notes
            self._refresh_notes_based_popups()

    def _on_extracted_data_changed(self, data: dict):
        """Handle extracted data updates from SharedDataStore - full popup population."""
        if not data:
            return

        # Skip if report data exists (report takes priority over notes)
        if self._has_report_data():
            print(f"[TRIBUNAL] Skipping extracted data from SharedDataStore - report data already imported")
            return

        print(f"[TRIBUNAL] Received extracted data from SharedDataStore: {list(data.keys())}")

        # Get categories from the data structure
        categories = data.get("categories", data)
        if not categories:
            return

        # Store at page level
        if not hasattr(self, '_extracted_categories'):
            self._extracted_categories = {}
        self._extracted_categories = categories

        # Delete cached popups so they get recreated with new data
        popups_to_refresh = ['forensic', 'previous_mh_dates', 'current_admission', 'diagnosis',
                            'treatment', 'strengths', 'risk_harm', 'risk_property', 'progress']
        for key in popups_to_refresh:
            if key in self.popups:
                old_popup = self.popups[key]
                self.popup_stack.removeWidget(old_popup)
                old_popup.deleteLater()
                del self.popups[key]
                print(f"[TRIBUNAL] Deleted cached '{key}' popup for refresh")

        print(f"[TRIBUNAL] Refreshed popups with SharedDataStore data")

    def _refresh_notes_based_popups(self):
        """Refresh popups that depend on notes data."""
        notes_popups = ['risk_harm', 'risk_property', 'previous_mh_dates', 'progress']
        for key in notes_popups:
            if key in self.popups:
                old_popup = self.popups[key]
                self.popup_stack.removeWidget(old_popup)
                old_popup.deleteLater()
                del self.popups[key]
                print(f"[TRIBUNAL] Deleted cached '{key}' popup for notes refresh")

    def _on_report_sections_changed(self, sections: dict, source_form: str):
        """Handle report sections imported from another form (cross-talk)."""
        # Only process if from nursing form (not our own import)
        if source_form == "tribunal":
            return

        print(f"[TRIBUNAL] Cross-talk received from {source_form}: {len(sections)} sections")

        # Store imported data for popups to use
        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        # Store imported data for popups — do NOT auto-fill cards
        for key, content in sections.items():
            if key in self.cards and content:
                # Store the imported data for the popup's imported data section
                self._imported_report_data[key] = content
                print(f"[TRIBUNAL] Cross-talk stored: {key}")

                # If popup already exists, populate it
                if key in self.popups:
                    self._populate_single_popup(self.popups[key], key, content)

        print(f"[TRIBUNAL] Cross-talk stored {len(sections)} sections from {source_form} (cards not auto-filled)")

    def set_notes(self, notes: list):
        """
        Set notes from shared data store.
        Called by MainWindow when notes are available from another section.
        """
        if not notes:
            return

        # Skip if report data exists (report takes priority over notes)
        if self._has_report_data():
            print(f"[Tribunal] Skipping set_notes - report data already imported")
            return

        # Skip if these exact notes were already processed
        notes_sig = (len(notes), id(notes))
        if self._notes_processed_id == notes_sig:
            print(f"[Tribunal] Skipping set_notes - notes already processed")
            return
        self._notes_processed_id = notes_sig

        # Store raw notes at page level for use in sections
        self._extracted_raw_notes = notes

        # If data extractor exists, update its notes too
        if hasattr(self, '_data_extractor') and self._data_extractor:
            if hasattr(self._data_extractor, 'set_notes'):
                self._data_extractor.set_notes(notes)

        print(f"[Tribunal] Received {len(notes)} notes from shared store")

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header bar
        header = QFrame()
        header.setFixedHeight(48)  # Reduced by 20%
        header.setStyleSheet("""
            QFrame {
                background: #7c3aed;
                border-bottom: 1px solid #6d28d9;
            }
        """)
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        # Back button
        back_btn = QPushButton("< Back")
        back_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: white;
                border: 1px solid rgba(255,255,255,0.3);
                padding: 6px 12px;
                border-radius: 6px;
                font-weight: 500;
            }
            QPushButton:hover {
                background: rgba(255,255,255,0.1);
            }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Psychiatric Tribunal Report")
        title.setStyleSheet("""
            font-size: 20px;
            font-weight: 700;
            color: white;
        """)
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Clear Report button on the right
        clear_btn = QPushButton("Clear Report - Start New")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 16px;
            }
            QPushButton:hover {
                background: #7f1d1d;
            }
            QPushButton:pressed {
                background: #450a0a;
            }
        """)
        clear_btn.clicked.connect(self._clear_report)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = TribunalToolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        # Connect uploaded docs menu to SharedDataStore
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())
        # Track last active editor (persists when toolbar clicked)
        self._active_editor = None

        def cur():
            return self._active_editor

        def safe(method_name):
            editor = cur()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        # Connect formatting signals
        self.toolbar.set_font_family.connect(
            lambda family: cur().set_font_family(family) if cur() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: cur().set_font_size(size) if cur() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: cur().set_text_color(c) if cur() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: cur().set_highlight_color(c) if cur() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe("numbered_list"))
        self.toolbar.indent.connect(lambda: safe("indent"))
        self.toolbar.outdent.connect(lambda: safe("outdent"))
        self.toolbar.undo.connect(lambda: safe("undo"))
        self.toolbar.redo.connect(lambda: safe("redo"))
        self.toolbar.insert_date.connect(lambda: safe("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe("insert_section_break"))

        def check_spelling():
            editor = cur()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    from PySide6.QtWidgets import QMessageBox
                    QMessageBox.information(
                        self,
                        "Spell Check",
                        "No spelling errors found."
                    )

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area with splitter
        content = QWidget()
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(0, 0, 0, 0)
        content_layout.setSpacing(0)

        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(6)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #e5e7eb, stop:0.5 #9ca3af, stop:1 #e5e7eb);
                border-radius: 3px;
                margin: 40px 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #d1d5db, stop:0.5 #6b7280, stop:1 #d1d5db);
            }
        """)
        content_layout.addWidget(self.main_splitter)

        # Left: Cards
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setMinimumWidth(300)
        self.cards_holder.setStyleSheet("""
            QScrollArea {
                background: #f3f4f6;
                border: none;
            }
        """)
        self.main_splitter.addWidget(self.cards_holder)

        self.editor_root = QWidget()
        self.editor_root.setStyleSheet("background: #f3f4f6;")
        self.editor_root.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.editor_layout = QVBoxLayout(self.editor_root)
        self.editor_layout.setContentsMargins(32, 24, 32, 24)
        self.editor_layout.setSpacing(16)
        self.cards_holder.setWidget(self.editor_root)

        # Right: Panel (resizable)
        self.editor_panel = QFrame()
        self.editor_panel.setMinimumWidth(350)
        self.editor_panel.setMaximumWidth(800)
        self.editor_panel.setStyleSheet("""
            QFrame {
                background: rgba(245,245,245,0.98);
                border-left: 1px solid rgba(0,0,0,0.08);
            }
        """)
        self.main_splitter.addWidget(self.editor_panel)
        self.main_splitter.setStretchFactor(0, 1)
        self.main_splitter.setStretchFactor(1, 0)
        self.main_splitter.setSizes([600, 450])
        self.main_splitter.setCollapsible(0, False)
        self.main_splitter.setCollapsible(1, False)

        # Connect splitter movement to update card widths
        self.main_splitter.splitterMoved.connect(self._on_splitter_moved)

        panel_layout = QVBoxLayout(self.editor_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        # Header row with title and lock button
        self.panel_header = QWidget()
        self.panel_header.setStyleSheet("""
            background: rgba(139, 92, 246, 0.15);
            border-radius: 8px;
        """)
        header_layout = QHBoxLayout(self.panel_header)
        header_layout.setContentsMargins(12, 8, 12, 8)
        header_layout.setSpacing(8)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setWordWrap(True)
        self.panel_title.setStyleSheet("""
            font-size: 18px;
            font-weight: 700;
            color: #7c3aed;
            background: transparent;
        """)
        header_layout.addWidget(self.panel_title, 1)

        # Lock button in header
        self.header_lock_btn = QPushButton("Unlocked")
        self.header_lock_btn.setFixedSize(70, 26)
        self.header_lock_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.header_lock_btn.setToolTip("Click to lock this section")
        self.header_lock_btn.setStyleSheet("""
            QPushButton {
                background: rgba(34, 197, 94, 0.3);
                border: 2px solid #22c55e;
                border-radius: 13px;
                font-size: 13px;
                font-weight: 600;
                color: #16a34a;
            }
            QPushButton:hover { background: rgba(34, 197, 94, 0.5); }
        """)
        self.header_lock_btn.clicked.connect(self._toggle_current_popup_lock)
        self.header_lock_btn.hide()
        header_layout.addWidget(self.header_lock_btn)

        panel_layout.addWidget(self.panel_header)

        # Popup stack
        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        main_layout.addWidget(content)

        # Create all cards
        self._create_cards()

    # Sections that should be headings (no editor card, just clickable title)
    HEADING_ONLY_SECTIONS = {
        "learning_disability",     # Section 10
        "detention_required",      # Section 11
        "s2_detention",            # Section 19
        "other_detention",         # Section 20
    }

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self  # Capture reference to self for closure

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    def _toggle_current_popup_lock(self):
        """Toggle lock on the currently active popup."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'toggle_lock'):
            popup.toggle_lock()
            self._update_header_lock_button()

    def _update_header_lock_button(self):
        """Update header lock button to match current popup state."""
        popup = getattr(self, '_current_popup', None)
        if popup and hasattr(popup, 'is_locked'):
            is_locked = popup.is_locked()
            if is_locked:
                self.header_lock_btn.setText("Locked")
                self.header_lock_btn.setToolTip("Click to unlock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(239, 68, 68, 0.3);
                        border: 2px solid #ef4444;
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: #dc2626;
                    }
                    QPushButton:hover { background: rgba(239, 68, 68, 0.5); }
                """)
            else:
                self.header_lock_btn.setText("Unlocked")
                self.header_lock_btn.setToolTip("Click to lock this section")
                self.header_lock_btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(34, 197, 94, 0.3);
                        border: 2px solid #22c55e;
                        border-radius: 13px;
                        font-size: 13px;
                        font-weight: 600;
                        color: #16a34a;
                    }
                    QPushButton:hover { background: rgba(34, 197, 94, 0.5); }
                """)
            self.header_lock_btn.show()
        else:
            self.header_lock_btn.hide()

    def _set_current_popup(self, popup):
        """Set the current active popup and update lock button."""
        self._current_popup = popup
        self._update_header_lock_button()

    def _create_cards(self):
        """Create all section cards (or headings for certain sections)."""
        for title, key in self.SECTIONS:
            if key in self.HEADING_ONLY_SECTIONS:
                # Create heading widget instead of full card
                card = TribunalHeadingWidget(title, key, parent=self.editor_root)
            else:
                # Create standard card widget
                card = TribunalCardWidget(title, key, parent=self.editor_root)
                # Hook up focus event to register this editor as active
                self._hook_editor_focus(card.editor)
                # Connect card text changes to sync with popup
                card.editor.textChanged.connect(lambda k=key: self._on_card_text_changed(k))
            card.clicked.connect(self._on_card_clicked)
            self.cards[key] = card
            self.editor_layout.addWidget(card)

        self.editor_layout.addStretch()

    def _on_splitter_moved(self, pos, index):
        """Update card widths when splitter is moved."""
        self._update_card_widths()

    def _update_card_widths(self):
        """Update card widths based on current viewport size."""
        viewport_width = self.cards_holder.viewport().width()
        # Account for margins (32px on each side from editor_layout)
        card_width = viewport_width - 64
        if card_width > 100:  # Minimum sensible width
            for key, card in self.cards.items():
                if key in self.HEADING_ONLY_SECTIONS:
                    # Headings use max width so text wraps, but can be narrower
                    card.setMaximumWidth(card_width)
                else:
                    card.setFixedWidth(card_width)

    def showEvent(self, event):
        """Set initial card widths when page is shown."""
        super().showEvent(event)
        from PySide6.QtCore import QTimer
        QTimer.singleShot(50, self._update_card_widths)

        # Check for existing report sections in shared store (cross-talk)
        QTimer.singleShot(100, self._check_shared_store_for_sections)

    def _check_shared_store_for_sections(self):
        """Check SharedDataStore for existing sections when form is shown."""
        try:
            from shared_data_store import get_shared_store
            shared_store = get_shared_store()
            existing_sections = shared_store.report_sections
            source = shared_store.get_report_source()
            if existing_sections and source and source != "tribunal":
                print(f"[TRIBUNAL] showEvent: Found sections from {source}")
                self._on_report_sections_changed(existing_sections, source)
        except Exception as e:
            print(f"[TRIBUNAL] Error checking shared store: {e}")

    def resizeEvent(self, event):
        """Update card widths on window resize."""
        super().resizeEvent(event)
        self._update_card_widths()

    def _on_card_clicked(self, key: str):
        """Handle card click - show appropriate popup."""
        # Find section title
        title = next((t for t, k in self.SECTIONS if k == key), key)
        self.panel_title.setText(title)

        print(f"[TRIBUNAL] Card clicked: {key}")

        # Update card selection highlighting
        if self._selected_card_key and self._selected_card_key in self.cards:
            self.cards[self._selected_card_key].setSelected(False)
        if key in self.cards:
            self.cards[key].setSelected(True)
        self._selected_card_key = key

        # Create popup if not exists
        if key not in self.popups:
            popup = self._create_popup(key)
            if popup:
                self.popups[key] = popup
                self.popup_stack.addWidget(popup)
                # Handle different signal signatures
                if key == "forensic":
                    popup.sent.connect(lambda text, state, k=key: self._update_card(k, text))
                elif key in ("current_admission", "previous_mh_dates", "previous_admission_reasons", "progress", "risk_harm", "risk_property", "discharge_risk", "recommendations"):
                    # FixedDataPanel and RecentAdmissionPopup use sent signal with just text
                    popup.sent.connect(lambda text, k=key: self._update_card(k, text))
                else:
                    popup.sent.connect(lambda text, k=key: self._update_card(k, text))

                # Connect gender_changed signal from patient_details popup
                if key == "patient_details":
                    popup.gender_changed.connect(self._on_gender_changed)

                # For these popups, trigger initial send to populate card with pre-filled data
                if key in ("patient_details", "author", "signature") and hasattr(popup, '_send_to_card'):
                    popup._send_to_card()

                # Populate fixed panels with stored data when created (notes pipeline only)
                if key in ("previous_mh_dates", "previous_admission_reasons", "current_admission", "progress"):
                    if self._extracted_raw_notes and not self._has_report_data():
                        print(f"[TRIBUNAL] Populating newly created popup '{key}' with stored data")
                        self._populate_fixed_panels()

                # For incident panels (17 & 18), populate directly with incident data
                if key in ("risk_harm", "risk_property"):
                    if self._extracted_raw_notes and not self._has_report_data() and hasattr(self, '_incident_data') and self._incident_data:
                        print(f"[TRIBUNAL] Directly populating incident popup '{key}' with {len(self._incident_data)} incidents")
                        popup.set_entries(self._incident_data, f"{len(self._incident_data)} incidents")
                    elif self._extracted_raw_notes and not self._has_report_data():
                        print(f"[TRIBUNAL] Populating '{key}' via _populate_fixed_panels")
                        self._populate_fixed_panels()

                # For discharge_risk (section 21), use GPRRiskPopup with notes for risk analysis
                if key == "discharge_risk":
                    if self._extracted_raw_notes and not self._has_report_data():
                        print(f"[TRIBUNAL] Populating discharge_risk with {len(self._extracted_raw_notes)} notes for risk analysis")
                        popup.set_notes_for_risk_analysis(self._extracted_raw_notes)

                # For compliance (section 15), populate with all raw notes for non-compliance search
                if key == "compliance" and self._extracted_raw_notes and not self._has_report_data():
                    if hasattr(popup, 'set_notes'):
                        popup.set_notes(self._extracted_raw_notes)
                        print(f"[TRIBUNAL] Compliance popup searching {len(self._extracted_raw_notes)} notes")

                # For forensic popup - populate forensic history panel with notes analysis + extracted data
                # Skip if report data is imported (notes-based forensic data not needed)
                if key == "forensic" and not self._has_report_data():
                    if hasattr(self, '_pending_forensic_data'):
                        forensic_entries = self._pending_forensic_data.get('forensic', [])
                        forensic_notes = self._pending_forensic_data.get('forensic_notes', [])
                        if hasattr(popup, 'set_forensic_data'):
                            popup.set_forensic_data(forensic_notes, forensic_entries)
                            print(f"[TRIBUNAL] Populated forensic panel with notes analysis + {len(forensic_entries)} entries")
                        elif hasattr(popup, 'set_entries') and forensic_entries:
                            popup.set_entries(forensic_entries, f"{len(forensic_entries)} entries")
                            print(f"[TRIBUNAL] Populated forensic panel with {len(forensic_entries)} entries")

                # Populate imported report data into popup (if not already added at import time)
                if hasattr(self, '_imported_report_data') and key in self._imported_report_data:
                    if not getattr(popup, '_imported_data_added', False):
                        content = self._imported_report_data[key]
                        self._populate_single_popup(popup, key, content)

                # Special handling for discharge_risk - ensure imported data section exists
                if key == "discharge_risk" and hasattr(self, '_imported_report_data') and self._imported_report_data:
                    if not getattr(popup, '_imported_data_added', False):
                        if key in self._imported_report_data and self._imported_report_data[key]:
                            self._add_imported_data_to_popup(popup, key, self._imported_report_data[key])
                            print(f"[TRIBUNAL] Added imported data section for discharge_risk on click")

        # Show the popup
        if key in self.popups:
            self.popup_stack.setCurrentWidget(self.popups[key])
            self._set_current_popup(self.popups[key])

            # Send popup form content to card (imported data checkbox is unchecked so only form data flows)
            popup = self.popups[key]
            if hasattr(self, '_imported_report_data') and key in self._imported_report_data:
                if hasattr(popup, '_send_to_card'):
                    popup._send_to_card()

            # Legacy check for pre-extracted data (keeping for backward compatibility)
            if key in ("previous_admission_reasons", "current_admission", "progress", "risk_harm", "risk_property"):
                print(f"[TRIBUNAL] Checking for pre-extracted data for section '{key}'")

                # Try to get data from _extracted_categories first
                categories_to_use = None
                if hasattr(self, '_extracted_categories') and self._extracted_categories:
                    categories_to_use = self._extracted_categories
                    print(f"[TRIBUNAL] Using stored _extracted_categories")
                else:
                    # Fallback: get data directly from section 7 popup if it has been loaded
                    if "previous_admission_reasons" in self.popups:
                        sec7_popup = self.popups["previous_admission_reasons"]
                        if hasattr(sec7_popup, '_latest_panel_data') and sec7_popup._latest_panel_data:
                            panel_data = sec7_popup._latest_panel_data
                            categories_to_use = panel_data.get("categories", {})
                            if categories_to_use:
                                # Store it for future use
                                self._extracted_categories = categories_to_use
                                print(f"[TRIBUNAL] Got data from section 7 popup: {len(categories_to_use)} categories")

                # For section 8 (current_admission), get raw notes DIRECTLY from stored data
                # Skip if report data was imported (notes pipeline not needed)
                if key == "current_admission" and not (hasattr(self, '_imported_report_data') and 'current_admission' in self._imported_report_data):
                    popup = self.popups[key]

                    # Use the FULL raw notes stored at page level (14185 notes, not categorized 6)
                    raw_notes = self._extracted_raw_notes
                    print(f"[TRIBUNAL] Using {len(raw_notes)} raw notes from page-level storage")

                    if raw_notes:
                        # Filter notes by admission date
                        # 2 days before admission, 2 weeks after
                        filtered_notes, date_info = self._filter_raw_notes_around_admission(
                            raw_notes,
                            days_before=2,
                            days_after=14
                        )
                        # Convert notes to entry format
                        entries = []
                        for note in filtered_notes:
                            entry = {
                                "text": note.get("text") or note.get("content") or "",
                                "date": note.get("date", ""),
                            }
                            if entry["text"]:
                                entries.append(entry)

                        popup.set_entries(entries, date_info)
                        print(f"[TRIBUNAL] ✓ Set {len(entries)} entries from raw notes")
                    else:
                        print(f"[TRIBUNAL] No raw notes available for section 8")
                        popup.set_entries([], "Import notes first using the Import File button")

            # Pre-fill diagnosis section from raw notes
            if key == "diagnosis":
                self._prefill_diagnosis_from_notes()

            # Pre-fill learning disability section
            if key == "learning_disability":
                self._prefill_learning_disability()

            # Pre-fill medications in treatment section
            if key == "treatment":
                self._prefill_medications_from_notes()

    def _prefill_diagnosis_from_notes(self):
        """Extract diagnoses from ALL raw notes and pre-fill section 9."""
        import re
        from collections import Counter

        if "diagnosis" not in self.popups:
            return

        popup = self.popups["diagnosis"]

        # Check if already filled
        if any(combo.currentIndex() > 0 for combo in popup.dx_boxes):
            print(f"[TRIBUNAL] Diagnosis already filled - skipping prefill")
            return

        # Get ALL raw notes from page-level storage (not just section 7)
        raw_notes = self._extracted_raw_notes
        if not raw_notes:
            print(f"[TRIBUNAL] No raw notes for diagnosis prefill")
            return

        print(f"[TRIBUNAL] Searching {len(raw_notes)} notes for diagnoses...")

        # Patterns to find diagnosis mentions
        DIAGNOSIS_PATTERNS = [
            r'diagnosed\s+with[:\s]+([^\n\.]{5,100})',
            r'diagnosed\s+as[:\s]+([^\n\.]{5,100})',
            r'diagnosis\s+of[:\s]+([^\n\.]{5,100})',
            r'diagnosis[:\s]+([^\n\.]{5,100})',
            r'diagnosis\s*-\s*([^\n\.]{5,100})',
        ]

        # Substance types for M&BD matching
        SUBSTANCE_TYPES = [
            'cannabis', 'cannabinoid', 'cannabinoids',
            'alcohol',
            'opioid', 'opioids', 'opiate', 'opiates',
            'cocaine',
            'stimulant', 'stimulants', 'amphetamine', 'amphetamines',
            'sedative', 'sedatives', 'hypnotic', 'hypnotics', 'benzodiazepine',
            'hallucinogen', 'hallucinogens', 'lsd',
            'tobacco', 'nicotine',
            'volatile', 'solvent', 'solvents', 'inhalant',
        ]

        # Count all diagnosis mentions, track F-codes and M&BD substances
        diagnosis_counts = Counter()
        fcode_map = {}  # F-code -> diagnosis text
        mbd_substances = Counter()  # Track substance mentions for M&BD

        for note in raw_notes:
            text = note.get("text") or note.get("content") or ""
            text_lower = text.lower()

            # Check for M&BD / Mental and Behavioural Disorder patterns
            mbd_patterns = [
                r'mental\s+and\s+behaviou?ral\s+disorder[s]?\s+(?:due\s+to\s+(?:use\s+of\s+)?)?(\w+)',
                r'm\s*&\s*bd\s*[-:]\s*(\w+)',
                r'substance\s+(?:use|misuse|abuse)\s+disorder[s]?\s*[-:]?\s*(\w+)?',
            ]
            for pattern in mbd_patterns:
                matches = re.findall(pattern, text_lower)
                for match in matches:
                    if match:
                        # Check if the matched word is a substance
                        for substance in SUBSTANCE_TYPES:
                            if substance in match or match in substance:
                                mbd_substances[substance.split()[0]] += 1
                                break

            # Also check for standalone substance disorder mentions
            for substance in SUBSTANCE_TYPES:
                if f'{substance} dependence' in text_lower or f'{substance} use disorder' in text_lower:
                    mbd_substances[substance.split()[0]] += 1

            for pattern in DIAGNOSIS_PATTERNS:
                matches = re.findall(pattern, text, re.IGNORECASE)
                for match in matches:
                    match = match.strip()
                    # Split on common separators
                    parts = re.split(r'[,;]|\band\b', match)
                    for part in parts:
                        part = part.strip()

                        # Extract F-code if present (e.g., F25, F20.0)
                        fcode_match = re.search(r'\(?(F\d+\.?\d*)\)?', part, re.IGNORECASE)
                        if fcode_match:
                            fcode = fcode_match.group(1).upper()
                            # Store the diagnosis text associated with this F-code
                            clean_part = re.sub(r'\([^)]*\)', '', part).strip(' .-:').lower()
                            if clean_part and len(clean_part) > 3:
                                fcode_map[fcode] = clean_part

                        # Remove ICD codes in parentheses for counting
                        part = re.sub(r'\([^)]*\)', '', part)
                        part = part.strip(' .-:')

                        # Filter out noise
                        if len(part) > 5 and len(part) < 100:
                            normalized = part.lower().strip()
                            # Skip common false positives
                            if normalized not in ['the', 'and', 'with', 'for', 'has', 'was', 'been', 'being', 'that', 'this']:
                                diagnosis_counts[normalized] += 1

        if not diagnosis_counts and not mbd_substances:
            print(f"[TRIBUNAL] No diagnoses found in notes")
            return

        # Get most common diagnoses
        most_common = diagnosis_counts.most_common(15)
        print(f"[TRIBUNAL] Found diagnoses (by frequency): {most_common}")
        if fcode_map:
            print(f"[TRIBUNAL] Found F-codes: {fcode_map}")
        if mbd_substances:
            print(f"[TRIBUNAL] Found M&BD substances: {mbd_substances.most_common()}")

        # Build ICD-10 lookup from first combo
        icd10_lookup = {}  # normalized name -> (index, full_text)
        fcode_lookup = {}  # F-code -> (index, full_text)
        first_combo = popup.dx_boxes[0]

        for i in range(1, first_combo.count()):
            item_text = first_combo.itemText(i)
            # Extract diagnosis name without ICD code
            name = item_text.split('(')[0].strip().lower()
            icd10_lookup[name] = (i, item_text)

            # Also index by F-code
            fcode_match = re.search(r'\((F\d+\.?\d*)\)', item_text, re.IGNORECASE)
            if fcode_match:
                fcode_lookup[fcode_match.group(1).upper()] = (i, item_text)

        # Match found diagnoses to ICD-10 entries
        matched_indices = []
        matched_texts = []
        matched_diagnosis_types = set()  # Track matched diagnosis types to avoid duplicates

        # First, try matching by F-codes (most reliable)
        for fcode, dx_text in fcode_map.items():
            if len(matched_indices) >= 3:
                break
            if fcode in fcode_lookup:
                idx, text = fcode_lookup[fcode]
                if idx not in matched_indices:
                    matched_indices.append(idx)
                    matched_texts.append(text)
                    # Track the diagnosis type
                    dx_type = text.split('(')[0].strip().lower()
                    matched_diagnosis_types.add(dx_type)
                    print(f"[TRIBUNAL]   F-code match: {fcode} -> {text}")

        # Normalize M&BD substances (cannabis/cannabinoid/cannabinoids are all the same)
        mbd_substance_map = {
            'cannabis': 'cannabinoids',
            'cannabinoid': 'cannabinoids',
            'cannabinoids': 'cannabinoids',
            'alcohol': 'alcohol',
            'opioid': 'opioids',
            'opioids': 'opioids',
            'opiate': 'opioids',
            'opiates': 'opioids',
            'cocaine': 'cocaine',
            'stimulant': 'other stimulants',
            'stimulants': 'other stimulants',
            'amphetamine': 'other stimulants',
            'amphetamines': 'other stimulants',
            'sedative': 'sedatives or hypnotics',
            'sedatives': 'sedatives or hypnotics',
            'hypnotic': 'sedatives or hypnotics',
            'hypnotics': 'sedatives or hypnotics',
            'benzodiazepine': 'sedatives or hypnotics',
            'hallucinogen': 'hallucinogens',
            'hallucinogens': 'hallucinogens',
            'tobacco': 'tobacco',
            'nicotine': 'tobacco',
            'volatile': 'volatile solvents',
            'solvent': 'volatile solvents',
            'solvents': 'volatile solvents',
        }

        # Deduplicate M&BD substances (combine cannabis+cannabinoid+cannabinoids counts)
        normalized_mbd = Counter()
        for substance, sub_count in mbd_substances.items():
            normalized = mbd_substance_map.get(substance, substance)
            normalized_mbd[normalized] += sub_count
        print(f"[TRIBUNAL] Normalized M&BD substances: {normalized_mbd.most_common()}")

        # PRIORITY ORDER:
        # 1. F-codes (already done above)
        # 2. Most common diagnoses from text (schizoaffective, etc.)
        # 3. M&BD substances (if not already matched)

        # Step 2: Match most common diagnoses by text
        for found_dx, count in most_common:
            if len(matched_indices) >= 3:
                break

            # Skip noise and M&BD entries (handle M&BD separately)
            skip_terms = [
                'worker', 'mental', 'practitioner', 'not mentioned',
                'behavioural disorder due to', 'm&bd', 'unspecified'
            ]
            if any(term in found_dx for term in skip_terms):
                continue

            # Skip duplicates like "of schizoaffective disorder" when we have "schizoaffective disorder"
            if found_dx.startswith('of ') or found_dx.startswith('his ') or found_dx.startswith('her '):
                continue

            # Check if this is a duplicate diagnosis type we've already matched
            # Extract base diagnosis type from found_dx
            def get_base_diagnosis_type(dx):
                """Extract the core diagnosis type to prevent duplicates."""
                dx = dx.lower()
                # Remove common prefixes
                for prefix in ['of ', 'his main diagnosis was of ', 'her main diagnosis was of ', 'the ']:
                    if dx.startswith(prefix):
                        dx = dx[len(prefix):]
                # Extract core type
                core_types = [
                    'schizoaffective', 'schizophrenia', 'bipolar', 'depression', 'depressive',
                    'personality disorder', 'anxiety', 'ptsd', 'autism', 'adhd', 'psychotic',
                    'mania', 'manic'
                ]
                for ct in core_types:
                    if ct in dx:
                        return ct
                return dx[:20]  # First 20 chars as fallback

            found_base_type = get_base_diagnosis_type(found_dx)
            if found_base_type in matched_diagnosis_types:
                print(f"[TRIBUNAL]   Skipping duplicate type '{found_base_type}': '{found_dx}'")
                continue

            # Try exact match first
            if found_dx in icd10_lookup:
                idx, text = icd10_lookup[found_dx]
                if idx not in matched_indices:
                    matched_indices.append(idx)
                    matched_texts.append(text)
                    matched_diagnosis_types.add(found_base_type)
                    print(f"[TRIBUNAL]   Exact match: '{found_dx}' -> {text} (count: {count})")
                continue

            # Try to find the best matching ICD entry
            best_match = None
            best_match_score = 0

            for icd_name, (idx, text) in icd10_lookup.items():
                if idx in matched_indices:
                    continue

                # Skip if icd_name is too short
                if len(icd_name) < 5:
                    continue

                # Calculate match score
                score = 0

                # High score: found diagnosis starts with ICD name or vice versa
                if found_dx.startswith(icd_name) or icd_name.startswith(found_dx):
                    score = 100

                # Medium score: one fully contains the other (must be substantial)
                elif len(found_dx) > 10 and len(icd_name) > 10:
                    if found_dx in icd_name:
                        score = 80
                    elif icd_name in found_dx:
                        score = 80

                # Lower score: key diagnostic terms match
                if score == 0:
                    # Use FULL diagnostic terms (not shortened)
                    DIAGNOSTIC_TERMS = [
                        'schizoaffective disorder',
                        'schizoaffective',
                        'schizophrenia',
                        'bipolar affective disorder',
                        'bipolar disorder',
                        'recurrent depressive disorder',
                        'depressive episode',
                        'dissocial personality disorder',
                        'dissocial personality',
                        'emotionally unstable personality',
                        'borderline personality',
                        'paranoid schizophrenia',
                        'paranoid personality',
                        'psychotic disorder',
                        'anxiety disorder',
                        'generalised anxiety',
                        'post-traumatic stress',
                        'ptsd',
                        'autism spectrum disorder',
                        'attention deficit',
                        'manic episode',
                        'hypomania',
                    ]

                    for term in DIAGNOSTIC_TERMS:
                        # Both must contain the FULL term
                        if term in found_dx and term in icd_name:
                            score = 70
                            break

                if score > best_match_score:
                    best_match_score = score
                    best_match = (idx, text)

            if best_match and best_match_score >= 70:
                idx, text = best_match
                matched_indices.append(idx)
                matched_texts.append(text)
                matched_diagnosis_types.add(found_base_type)
                print(f"[TRIBUNAL]   Best match (score {best_match_score}): '{found_dx}' -> {text} (count: {count})")

        # Step 3: Match M&BD substances (only if we still have slots available)
        if len(matched_indices) < 3 and normalized_mbd:
            # Check if we already have a cannabinoid/substance match
            has_substance = any('m&bd' in t.lower() for t in matched_texts)
            if not has_substance:
                for substance, sub_count in normalized_mbd.most_common(1):  # Only the top substance
                    if len(matched_indices) >= 3:
                        break
                    # Find matching M&BD entry in ICD-10
                    for icd_name, (idx, text) in icd10_lookup.items():
                        if idx in matched_indices:
                            continue
                        # Match M&BD entries with the substance
                        if 'm&bd' in icd_name and substance in icd_name:
                            # Prefer "Harmful use" entry (simpler)
                            if 'harmful use' in icd_name:
                                matched_indices.append(idx)
                                matched_texts.append(text)
                                dx_type = text.split('(')[0].strip().lower()
                                matched_diagnosis_types.add(dx_type)
                                print(f"[TRIBUNAL]   M&BD match: '{substance}' -> {text}")
                                break

        # Pre-fill only as many as we found (1, 2, or 3)
        for i, (idx, text) in enumerate(zip(matched_indices, matched_texts)):
            if i < len(popup.dx_boxes):
                popup.dx_boxes[i].blockSignals(True)
                popup.dx_boxes[i].setCurrentIndex(idx)
                popup.dx_boxes[i].blockSignals(False)

        if matched_indices:
            popup._update_preview()
            print(f"[TRIBUNAL] ✓ Pre-filled {len(matched_indices)} diagnosis(es)")

    def _prefill_learning_disability(self):
        """Check notes and diagnoses for learning disability and pre-fill section 10."""
        import re

        if "learning_disability" not in self.popups:
            return

        popup = self.popups["learning_disability"]

        # Get ALL raw notes from page-level storage
        raw_notes = self._extracted_raw_notes
        if not raw_notes:
            print(f"[TRIBUNAL] No raw notes for learning disability check")
            # Still set to No if no notes
            popup.q1_no_btn.setChecked(True)
            return

        print(f"[TRIBUNAL] Searching {len(raw_notes)} notes for learning disability...")

        ld_found = False

        # Step 1: Check if section 9 has an LD diagnosis (F70-F79 = Intellectual disabilities)
        if "diagnosis" in self.popups:
            dx_popup = self.popups["diagnosis"]
            for combo in dx_popup.dx_boxes:
                selected_text = combo.currentText().lower()
                # Check for F70-F79 codes or LD keywords in selected diagnosis
                if re.search(r'\(f7[0-9]', selected_text):
                    ld_found = True
                    print(f"[TRIBUNAL]   LD diagnosis found in section 9: {combo.currentText()}")
                    break
                if any(kw in selected_text for kw in ['learning disability', 'intellectual disability', 'mental retardation']):
                    ld_found = True
                    print(f"[TRIBUNAL]   LD diagnosis found in section 9: {combo.currentText()}")
                    break

        # Step 2: Search notes for learning disability mentions
        if not ld_found:
            ld_keywords = [
                'learning disability',
                'learning disabilities',
                'learning difficulties',
                'intellectual disability',
                'mental retardation',
                'ld diagnosis',
                'diagnosed with ld',
                'mild ld',
                'moderate ld',
                'severe ld',
                'profound ld',
                'borderline ld',
                'borderline intellectual',
            ]

            # Also check for F70-F79 codes in notes
            ld_fcode_pattern = r'\bF7[0-9]\.?\d*\b'

            for note in raw_notes:
                text = (note.get("text") or note.get("content") or "").lower()

                # Check for F-codes
                if re.search(ld_fcode_pattern, text, re.IGNORECASE):
                    ld_found = True
                    print(f"[TRIBUNAL]   Found LD F-code (F70-F79) in notes")
                    break

                # Check for keywords
                for kw in ld_keywords:
                    if kw in text:
                        # Check it's not negated (e.g., "no learning disability")
                        idx = text.find(kw)
                        # Check 30 chars before for negation
                        before = text[max(0, idx-30):idx]
                        if not any(neg in before for neg in ['no ', 'not ', 'without ', 'denies ', 'nil ', 'no evidence']):
                            ld_found = True
                            print(f"[TRIBUNAL]   Found: '{kw}' in notes")
                            break
                if ld_found:
                    break

        # Set the radio button - ALWAYS set one or the other
        if ld_found:
            popup.q1_yes_btn.setChecked(True)
            popup.q1_no_btn.setChecked(False)
            print(f"[TRIBUNAL] ✓ Learning disability FOUND - set to YES")
        else:
            popup.q1_no_btn.setChecked(True)
            popup.q1_yes_btn.setChecked(False)
            print(f"[TRIBUNAL] ✓ Learning disability NOT found - set to NO")

        # Update preview
        if hasattr(popup, '_update_preview'):
            popup._update_preview()

    def _prefill_medications_from_notes(self):
        """Extract medications from notes (last year only) and pre-fill section 12 with most recent per class."""
        import re
        from datetime import datetime, timedelta
        from CANONICAL_MEDS import MEDICATIONS

        if "treatment" not in self.popups:
            return

        popup = self.popups["treatment"]

        # Get ALL raw notes from page-level storage
        raw_notes = self._extracted_raw_notes
        if not raw_notes:
            print(f"[TRIBUNAL] No raw notes for medication extraction")
            return

        # === PARSE DATE HELPER ===
        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # === FIND MOST RECENT NOTE DATE AND FILTER TO LAST YEAR ===
        note_dates = [parse_date(n.get("date") or n.get("datetime")) for n in raw_notes]
        note_dates = [d for d in note_dates if d]
        if not note_dates:
            print(f"[TRIBUNAL] No dates found in notes")
            return

        latest_date = max(note_dates)
        cutoff_date = latest_date - timedelta(days=365)
        print(f"[TRIBUNAL] Latest note: {latest_date.date()}, searching from {cutoff_date.date()}")

        # Filter notes to last year only
        recent_notes = []
        for n in raw_notes:
            note_date = parse_date(n.get("date") or n.get("datetime"))
            if note_date and note_date >= cutoff_date:
                recent_notes.append(n)

        print(f"[TRIBUNAL] Extracting medications from {len(recent_notes)} notes (last year)...")

        # === BUILD TOKEN INDEX ===
        token_map = {}  # token -> (key, canonical)
        meta_map = {}   # key -> metadata

        for key, meta in MEDICATIONS.items():
            meta_map[key] = meta
            canonical = meta["canonical"]
            for syn in meta.get("patterns", []):
                s = syn.lower().strip()
                if s:
                    token_map[s] = (key, canonical)

        def tokenise(text):
            text = text.lower()
            text = re.sub(r'(?<=\d)\.(?=\d)', 'DOT', text)
            text = re.sub(r'[^a-z0-9DOT]+', ' ', text)
            text = text.replace("DOT", ".")
            return text.split()

        def find_dose(tokens, idx):
            unit_set = {"mg", "mcg", "µg", "g", "units", "iu"}
            for i in range(idx + 1, min(idx + 4, len(tokens))):
                tok = tokens[i]
                m = re.match(r'(\d+(?:\.\d+)?)(mg|mcg|µg|g|units|iu)$', tok)
                if m:
                    return float(m.group(1)), m.group(2)
                if tok.isdigit() or re.match(r'\d+\.\d+', tok):
                    if i + 1 < len(tokens) and tokens[i + 1] in unit_set:
                        return float(tok), tokens[i + 1]
            return None, None

        def find_freq(tokens, idx):
            freq_set = {"od", "bd", "tds", "qds", "qid", "nocte", "mane", "stat", "prn", "daily", "weekly", "monthly"}
            for i in range(idx, min(idx + 5, len(tokens))):
                if tokens[i] in freq_set:
                    return tokens[i]
            return None

        # === PSYCHIATRIC MEDICATION CLASSES (prioritized) ===
        PSYCH_CLASSES = {
            "Antipsychotic": 1,
            "Antidepressant": 2,
            "Mood Stabiliser": 3,
            "Anxiolytic": 4,
            "Benzodiazepine": 5,
            "Anticonvulsant / Benzodiazepine": 5,
            "Sedative": 6,
            "Sedation": 6,
            "Sleep": 7,
            "Stimulant": 8,
            "Addictions": 9,
            "Opioid Substitution": 9,
        }

        # === EXTRACT MEDICATIONS ===
        meds_found = []
        for n in recent_notes:
            content = n.get("content", "") or n.get("text", "") or ""
            if not content:
                continue
            tokens = tokenise(content)
            note_date = parse_date(n.get("date") or n.get("datetime"))

            for i, tok in enumerate(tokens):
                if tok in token_map:
                    key, canonical = token_map[tok]
                    meta = meta_map[key]
                    med_class = meta.get("class", "Other")

                    strength, unit = find_dose(tokens, i)
                    freq = find_freq(tokens, i)

                    meds_found.append({
                        "med_key": key,
                        "canonical": canonical,
                        "class": med_class,
                        "strength": strength,
                        "unit": unit or "mg",
                        "frequency": freq,
                        "date": note_date,
                    })

        print(f"[TRIBUNAL] Found {len(meds_found)} medication mentions in last year")

        if not meds_found:
            print(f"[TRIBUNAL] No medications found in notes")
            return

        # === GROUP BY CLASS AND PICK MOST RECENT PER CLASS ===
        class_groups = {}
        for med in meds_found:
            med_class = med.get("class", "Other")
            if med_class not in class_groups:
                class_groups[med_class] = []
            class_groups[med_class].append(med)

        # For each class, find the most recent medication with a dose
        most_recent_by_class = []
        for med_class, mentions in class_groups.items():
            # Sort by date descending
            sorted_mentions = sorted(
                mentions,
                key=lambda x: x.get("date") or datetime.min,
                reverse=True
            )
            # Take most recent with a dose
            for m in sorted_mentions:
                if m.get("strength"):
                    most_recent_by_class.append(m)
                    break
            else:
                # If none have dose, take most recent anyway
                if sorted_mentions:
                    most_recent_by_class.append(sorted_mentions[0])

        # === PRIORITIZE PSYCHIATRIC CLASSES ===
        def class_priority(med):
            med_class = med.get("class", "Other")
            return PSYCH_CLASSES.get(med_class, 100)  # Non-psych classes get low priority

        # Sort: psych classes first (by priority), then others
        most_recent_by_class.sort(key=class_priority)

        # Limit to top entries (psych meds first)
        max_meds = 10
        final_meds = most_recent_by_class[:max_meds]

        print(f"[TRIBUNAL] Selected {len(final_meds)} medications (1 per class, psych prioritized):")
        for m in final_meds:
            strength = m.get('strength')
            dose_str = f"{strength}{m.get('unit', 'mg')}" if strength else "no dose"
            print(f"[TRIBUNAL]   [{m.get('class')}] {m.get('canonical')}: {dose_str} {m.get('frequency') or ''} ({m.get('date').date() if m.get('date') else 'no date'})")

        # Map extracted frequency to popup frequency options
        FREQ_MAP = {
            "od": "OD", "daily": "OD", "once daily": "OD", "1x daily": "OD",
            "bd": "BD", "twice daily": "BD", "2x daily": "BD",
            "tds": "TDS", "three times daily": "TDS", "3x daily": "TDS",
            "qds": "QDS", "qid": "QDS", "four times daily": "QDS", "4x daily": "QDS",
            "nocte": "Nocte", "at night": "Nocte", "night": "Nocte", "on": "Nocte",
            "mane": "Mane", "in the morning": "Mane", "am": "Mane",
            "prn": "PRN", "as required": "PRN", "when required": "PRN",
            "weekly": "Weekly", "1 weekly": "Weekly", "once weekly": "Weekly",
            "fortnightly": "Fortnightly", "2 weekly": "Fortnightly", "every 2 weeks": "Fortnightly",
            "3 weekly": "3 Weekly", "every 3 weeks": "3 Weekly",
            "monthly": "Monthly", "4 weekly": "Monthly", "every 4 weeks": "Monthly",
        }

        # Populate the medication entries
        while len(popup._medications) < len(final_meds):
            popup._add_medication_entry()

        for i, med in enumerate(final_meds):
            if i >= len(popup._medications):
                break

            entry = popup._medications[i]
            med_key = med.get("med_key", "")

            # Set medication name (the key in MEDICATIONS dict)
            name_combo = entry.get("name")
            if name_combo:
                idx = name_combo.findText(med_key)
                if idx >= 0:
                    name_combo.setCurrentIndex(idx)
                else:
                    # Try canonical name
                    canonical = med.get("canonical", "")
                    idx = name_combo.findText(canonical.upper())
                    if idx >= 0:
                        name_combo.setCurrentIndex(idx)

            # Set dose
            dose_combo = entry.get("dose")
            if dose_combo and med.get("strength"):
                strength = med.get("strength")
                unit = med.get("unit", "mg")
                if isinstance(strength, float) and strength == int(strength):
                    dose_str = f"{int(strength)}{unit}"
                else:
                    dose_str = f"{strength}{unit}"
                idx = dose_combo.findText(dose_str)
                if idx >= 0:
                    dose_combo.setCurrentIndex(idx)
                else:
                    dose_combo.setCurrentText(dose_str)

            # Set frequency
            freq_combo = entry.get("freq")
            if freq_combo and med.get("frequency"):
                freq = med.get("frequency", "").lower()
                mapped_freq = FREQ_MAP.get(freq, "OD")
                idx = freq_combo.findText(mapped_freq)
                if idx >= 0:
                    freq_combo.setCurrentIndex(idx)

        # Update preview
        if hasattr(popup, '_update_preview'):
            popup._update_preview()

        print(f"[TRIBUNAL] ✓ Pre-filled {len(final_meds)} medication(s)")

    def _score_and_display_risk_notes(self, notes):
        """Score notes using riskDICT and display results in terminal."""
        from datetime import datetime
        from pathlib import Path

        # Load risk dictionary
        risk_dict = {}
        risk_file = Path(__file__).parent / "riskDICT.txt"
        if not risk_file.exists():
            print(f"[TRIBUNAL] riskDICT.txt not found at {risk_file}")
            return

        with open(risk_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line or ',' not in line:
                    continue
                parts = line.rsplit(',', 1)
                term = parts[0].strip().lower()
                score_str = parts[1].strip() if len(parts) > 1 else ''
                if term and score_str:
                    try:
                        score = int(score_str)
                        risk_dict[term] = score
                    except:
                        pass

        print(f"\n[TRIBUNAL] ========== RISK SCORING (riskDICT: {len(risk_dict)} terms) ==========")

        def parse_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        def score_note(note):
            content = (note.get('content', '') or note.get('text', '') or '').lower()
            if not content:
                return 0, []
            total_score = 0
            matches = []
            for term, score in risk_dict.items():
                if term in content:
                    total_score += score
                    matches.append((term, score))
            return total_score, matches

        # Score all notes
        scored_notes = []
        for note in notes:
            score, matches = score_note(note)
            note_date = parse_date(note.get('date') or note.get('datetime'))
            note_type = note.get('type', 'Unknown')
            preview = (note.get('content', '') or note.get('text', '') or '')[:100]
            scored_notes.append({
                'date': note_date,
                'type': note_type,
                'score': score,
                'matches': matches,
                'preview': preview,
            })

        # Sort by date (oldest first for chronological view)
        scored_notes.sort(key=lambda x: x['date'] or datetime.min)

        # Display ALL notes in date order
        print(f"\n[TRIBUNAL] === ALL {len(scored_notes)} NOTES - RISK SCORED (DATE ORDER) ===\n")
        print(f"{'#':>3} | {'DATE':<12} | {'SCORE':>6} | {'TYPE':<20} | PREVIEW")
        print("-" * 120)

        for i, n in enumerate(scored_notes):
            date_str = n['date'].strftime('%d/%m/%Y') if n['date'] else 'no date'
            preview = n['preview'][:50].replace('\n', ' ')
            print(f"{i+1:>3} | {date_str:<12} | {n['score']:>6} | {n['type']:<20} | {preview}...")

        # Score distribution
        high_risk = len([n for n in scored_notes if n['score'] >= 500])
        med_risk = len([n for n in scored_notes if 100 <= n['score'] < 500])
        low_risk = len([n for n in scored_notes if 0 < n['score'] < 100])
        negative = len([n for n in scored_notes if n['score'] < 0])
        neutral = len([n for n in scored_notes if n['score'] == 0])

        print(f"\n[TRIBUNAL] === SCORE DISTRIBUTION ===")
        print(f"  High risk (>=500):  {high_risk}")
        print(f"  Medium (100-499):   {med_risk}")
        print(f"  Low (1-99):         {low_risk}")
        print(f"  Neutral (0):        {neutral}")
        print(f"  Protective (<0):    {negative}")
        print(f"[TRIBUNAL] ================================================\n")

    def _on_gender_changed(self, gender: str):
        """Handle gender change from patient details popup."""
        self._current_gender = gender
        print(f"[TRIBUNAL] Gender changed to: {gender}")

        # Update FactorsHearingPopup if it exists
        if "factors_hearing" in self.popups:
            self.popups["factors_hearing"].set_gender(gender)
            self.popups["factors_hearing"]._update_preview()

        # Update AdjustmentsPopup if it exists
        if "adjustments" in self.popups:
            self.popups["adjustments"].set_gender(gender)

        # Update ForensicHistoryPopup if it exists
        if "forensic" in self.popups:
            self.popups["forensic"].update_gender(gender)

        # Update PastPsychPopup if it exists
        if "previous_mh_dates" in self.popups:
            self.popups["previous_mh_dates"].update_gender(gender)

        # Update TreatmentPopup if it exists
        if "treatment" in self.popups:
            self.popups["treatment"].set_gender(gender)

        # Update StrengthsPopup if it exists
        if "strengths" in self.popups:
            self.popups["strengths"].set_gender(gender)

        # Update CompliancePopup if it exists
        if "compliance" in self.popups:
            self.popups["compliance"].set_gender(gender)

        # Update DischargeRiskPopup if it exists
        if "discharge_risk" in self.popups:
            self.popups["discharge_risk"].set_gender(gender)

        # Update CommunityManagementPopup if it exists
        if "community" in self.popups:
            self.popups["community"].set_gender(gender)

        # Update RecommendationsPopup if it exists
        if "recommendations" in self.popups:
            self.popups["recommendations"].set_gender(gender)

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details popup and card from extracted demographics."""
        if not patient_info:
            return

        # Create popup if it doesn't exist
        if "patient_details" not in self.popups:
            from tribunal_popups import PatientDetailsPopup
            popup = PatientDetailsPopup(parent=self)
            self.popups["patient_details"] = popup
            self.popup_stack.addWidget(popup)
            # Connect the sent signal so edits update the card
            popup.sent.connect(lambda text: self._update_card("patient_details", text))
            popup.gender_changed.connect(self._on_gender_changed)
            print("[TRIBUNAL] Created and connected patient_details popup from _fill_patient_details")

        # Fill the popup fields
        popup = self.popups["patient_details"]
        if hasattr(popup, 'fill_patient_info'):
            popup.fill_patient_info(patient_info)

            # Also update the card preview
            text = popup.generate_text()
            if "patient_details" in self.cards and text.strip():
                self.cards["patient_details"].editor.setPlainText(text)
                print(f"[TribunalReport] Updated patient_details card with demographics")

    def _create_popup(self, key: str):
        """Create the appropriate popup for a section."""
        from tribunal_popups import AuthorPopup, PatientDetailsPopup, YesNoDetailsPopup

        if key == "author":
            return AuthorPopup(parent=self, my_details=self._my_details)
        elif key == "patient_details":
            return PatientDetailsPopup(parent=self)
        elif key == "factors_hearing":
            from tribunal_popups import FactorsHearingPopup
            return FactorsHearingPopup(parent=self, gender=self._current_gender)
        elif key == "adjustments":
            from tribunal_popups import AdjustmentsPopup
            return AdjustmentsPopup(parent=self, gender=self._current_gender)
        elif key == "forensic":
            from forensic_history_popup import ForensicHistoryPopup
            return ForensicHistoryPopup(parent=self, gender=self._current_gender, show_index_offence=True)
        elif key == "previous_mh_dates":
            # Section 6: Past psychiatric history (matches GPR section 6 exactly)
            from tribunal_popups import TribunalPsychHistoryPopup
            return TribunalPsychHistoryPopup(parent=self)
        elif key == "current_admission":
            # Section 8: Circumstances popup with preview and yellow entries (matches GPR section 3)
            from tribunal_popups import TribunalCircumstancesPopup
            return TribunalCircumstancesPopup(parent=self)
        elif key == "previous_admission_reasons":
            # Section 7: Same popup as section 6 (detected admissions with clerking notes)
            from tribunal_popups import TribunalPsychHistoryPopup
            return TribunalPsychHistoryPopup(parent=self)
        elif key == "progress":
            # Section 14: Progress popup with narrative summary and yellow entries
            # Filter to 1 year from most recent entry
            from tribunal_popups import TribunalProgressPopup
            return TribunalProgressPopup(parent=self, date_filter='1_year')
        elif key == "diagnosis":
            from tribunal_popups import DiagnosisPopup
            from icd10_dict import ICD10_DICT
            return DiagnosisPopup(icd10_dict=ICD10_DICT, parent=self)
        elif key == "learning_disability":
            from tribunal_popups import LearningDisabilityPopup
            return LearningDisabilityPopup(parent=self)
        elif key == "detention_required":
            from tribunal_popups import SimpleYesNoPopup
            return SimpleYesNoPopup(
                title="Mental Disorder Requiring Detention",
                question="Is there any mental disorder present which requires the patient to be detained in a hospital for assessment and/or medical treatment?",
                parent=self
            )
        elif key == "treatment":
            from tribunal_popups import TreatmentPopup
            return TreatmentPopup(parent=self, gender=self._current_gender)
        elif key == "strengths":
            from tribunal_popups import StrengthsPopup
            return StrengthsPopup(parent=self, gender=self._current_gender)
        elif key == "compliance":
            from tribunal_popups import TribunalCompliancePopup
            return TribunalCompliancePopup(parent=self, gender=self._current_gender)
        elif key == "mca_dol":
            from tribunal_popups import DoLsPopup
            return DoLsPopup(parent=self)
        elif key == "risk_harm":
            # Section 17: Risk harm popup with categorized incidents
            from tribunal_popups import TribunalRiskHarmPopup
            return TribunalRiskHarmPopup(parent=self)
        elif key == "risk_property":
            # Section 18: Property damage popup with categorized incidents
            from tribunal_popups import TribunalRiskPropertyPopup
            return TribunalRiskPropertyPopup(parent=self)
        elif key == "s2_detention":
            from tribunal_popups import SimpleYesNoPopup
            return SimpleYesNoPopup(
                title="Section 2 Detention",
                question="Is detention under Section 2 justified for health, safety or protection of others?",
                parent=self
            )
        elif key == "other_detention":
            from tribunal_popups import SimpleYesNoPopup
            return SimpleYesNoPopup(
                title="Other Detention Sections",
                question="Is medical treatment justified for health, safety or protection of others?",
                parent=self
            )
        elif key == "discharge_risk":
            # Use GPRRiskPopup identical to GPR section 7
            from general_psychiatric_report_page import GPRRiskPopup
            return GPRRiskPopup(parent=self)
        elif key == "community":
            from tribunal_popups import CommunityManagementPopup
            return CommunityManagementPopup(parent=self, gender=self._current_gender)
        elif key == "recommendations":
            try:
                from general_psychiatric_report_page import GPRLegalCriteriaPopup
                from icd10_dict import ICD10_DICT
                print(f"[TRIBUNAL] Creating GPRLegalCriteriaPopup: ICD10_DICT={len(ICD10_DICT)} entries, gender={self._current_gender}")
                popup = GPRLegalCriteriaPopup(parent=self, gender=self._current_gender, icd10_dict=ICD10_DICT)
                print(f"[TRIBUNAL] GPRLegalCriteriaPopup created successfully")
                return popup
            except Exception as e:
                print(f"[TRIBUNAL] ERROR creating recommendations popup: {e}")
                import traceback
                traceback.print_exc()
                # Fallback: show error label
                from PySide6.QtWidgets import QLabel
                err = QLabel(f"ERROR loading recommendations popup:\n{e}")
                err.setStyleSheet("color: red; font-size: 16px; padding: 20px;")
                err.setWordWrap(True)
                return err
        elif key == "signature":
            from tribunal_popups import SignaturePopup
            return SignaturePopup(parent=self, my_details=self._my_details)

        # TODO: Add more popups as needed
        return None

    def _update_card(self, key: str, text: str):
        """Update a card with text from popup - uses smart update to preserve user additions."""
        print(f"[TRIBUNAL] _update_card called: key='{key}', text length={len(text) if text else 0}")
        if key not in self.cards:
            print(f"[TRIBUNAL] WARNING: key '{key}' not in self.cards")
            return

        new_text = text.strip() if text else ""

        # Prose sections should just replace the text (not use smart update)
        # Smart update is for structured "Field: Value" format only
        prose_sections = {
            "factors_hearing", "adjustments", "forensic", "previous_mh_dates",
            "previous_admission_reasons", "current_admission", "diagnosis",
            "treatment", "strengths", "progress", "compliance", "mca_dol",
            "risk_harm", "risk_property", "s2_detention", "other_detention",
            "discharge_risk", "community", "recommendations", "signature"
        }

        self._syncing = True
        try:
            if key in prose_sections:
                # For prose sections, just replace the text
                self.cards[key].editor.setPlainText(new_text)
            else:
                # For structured sections (patient_details, author), use smart update
                self._smart_update_card(key, new_text)
        finally:
            self._syncing = False
        print(f"[TRIBUNAL] Updated card '{key}'")

    def _smart_update_card(self, key: str, new_text: str):
        """Update card text intelligently - update existing lines, add new lines, preserve custom content."""
        if key not in self.cards:
            return

        card = self.cards[key]
        current_text = card.editor.toPlainText()

        # If card is empty, just set the text
        if not current_text.strip():
            card.editor.setPlainText(new_text)
            return

        # Parse both texts into field dictionaries
        current_fields = self._parse_structured_text(current_text)
        new_fields = self._parse_structured_text(new_text)

        # Build updated text: update known fields, preserve unknown content
        result_lines = []
        used_fields = set()

        # Process current text line by line
        for line in current_text.split('\n'):
            line_stripped = line.strip()
            if not line_stripped:
                result_lines.append(line)
                continue

            # Check if this line matches a known field pattern
            field_name = self._extract_field_name(line)
            if field_name and field_name in new_fields:
                # Update this field with new value
                result_lines.append(f"{field_name}: {new_fields[field_name]}")
                used_fields.add(field_name)
            elif field_name and field_name in current_fields:
                # Keep existing field value (not in new text)
                result_lines.append(line)
                used_fields.add(field_name)
            else:
                # Custom content - preserve it
                result_lines.append(line)

        # Add any new fields that weren't in current text
        for field_name, value in new_fields.items():
            if field_name not in used_fields:
                result_lines.append(f"{field_name}: {value}")

        card.editor.setPlainText('\n'.join(result_lines))

    def _parse_structured_text(self, text: str) -> dict:
        """Parse text with 'Field: Value' format into a dictionary."""
        fields = {}
        for line in text.split('\n'):
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    field_name = parts[0].strip()
                    value = parts[1].strip()
                    if field_name and value:
                        fields[field_name] = value
        return fields

    def _extract_field_name(self, line: str) -> str:
        """Extract field name from a 'Field: Value' line."""
        if ':' in line:
            parts = line.split(':', 1)
            if len(parts) == 2:
                return parts[0].strip()
        return None

    def _on_card_text_changed(self, key: str):
        """Handle card text changes - sync to popup if not already syncing."""
        if self._syncing:
            return

        if key not in self.cards:
            return

        card_text = self.cards[key].editor.toPlainText()

        # Parse card text and update corresponding popup
        if key in self.popups:
            self._syncing = True
            try:
                self._update_popup_from_card(key, card_text)
            finally:
                self._syncing = False

    def _update_popup_from_card(self, key: str, card_text: str):
        """Update popup fields from card text."""
        if key not in self.popups:
            return

        popup = self.popups[key]
        fields = self._parse_structured_text(card_text)

        # Patient details popup
        if key == "patient_details":
            # Handle name field (could be 'Name', 'Full Name', etc.)
            name_value = fields.get('Full Name') or fields.get('Name') or ''
            if name_value and hasattr(popup, 'name_field'):
                popup.name_field.blockSignals(True)
                popup.name_field.setText(name_value)
                popup.name_field.blockSignals(False)

            # Handle DOB field
            dob_value = fields.get('Date of Birth') or fields.get('DOB') or ''
            if dob_value and hasattr(popup, 'dob_field'):
                popup.dob_field.blockSignals(True)
                # Try to parse the date
                from PySide6.QtCore import QDate
                for fmt in ["dd/MM/yyyy", "dd-MM-yyyy", "yyyy-MM-dd", "d MMMM yyyy"]:
                    parsed = QDate.fromString(dob_value.strip(), fmt)
                    if parsed.isValid():
                        popup.dob_field.setDate(parsed)
                        break
                popup.dob_field.blockSignals(False)

            # Handle Gender field
            gender_value = fields.get('Gender') or ''
            if gender_value and hasattr(popup, 'gender_field'):
                popup.gender_field.blockSignals(True)
                idx = popup.gender_field.findText(gender_value, Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    popup.gender_field.setCurrentIndex(idx)
                popup.gender_field.blockSignals(False)
                # Propagate gender since signals were blocked
                self._on_gender_changed(gender_value)

            # Handle Residence field
            residence_value = fields.get('Usual Place of Residence') or fields.get('Residence') or fields.get('Address') or ''
            if residence_value and hasattr(popup, 'residence_field'):
                popup.residence_field.blockSignals(True)
                popup.residence_field.setPlainText(residence_value)
                popup.residence_field.blockSignals(False)

    def _advance_to_next_section(self, current_key: str):
        """Move to the next section after current one."""
        # Find current index in SECTIONS
        current_idx = None
        for idx, (title, key) in enumerate(self.SECTIONS):
            if key == current_key:
                current_idx = idx
                break

        if current_idx is None:
            return

        # Get next section (if exists)
        next_idx = current_idx + 1
        if next_idx < len(self.SECTIONS):
            next_title, next_key = self.SECTIONS[next_idx]
            print(f"[TRIBUNAL] Advancing to next section: {next_key}")
            self._on_card_clicked(next_key)

    def _on_data_extracted(self, key: str, data: dict):
        """Handle data extracted from DataExtractorPopup - fills sections with their SPECIFIC relevant category."""
        print(f"\n{'='*60}")
        print(f"[TRIBUNAL] _on_data_extracted called for section: '{key}'")
        print(f"[TRIBUNAL] Raw data keys: {list(data.keys())}")

        # Data format is {"categories": {"CATEGORY_NAME": {"items": [{"text": "..."}]}}}
        categories = data.get("categories", {})
        if not categories:
            categories = data

        print(f"[TRIBUNAL] Categories found: {list(categories.keys())}")
        for cat_name, cat_data in categories.items():
            if isinstance(cat_data, dict) and "items" in cat_data:
                print(f"[TRIBUNAL]   - {cat_name}: {len(cat_data.get('items', []))} items")
            else:
                print(f"[TRIBUNAL]   - {cat_name}: {type(cat_data)}")

        # Map each tribunal section to its PRIMARY category (one-to-one)
        SECTION_TO_CATEGORY = {
            "previous_mh_dates": "PAST_PSYCH",
            "previous_admission_reasons": "PAST_PSYCH",
            "current_admission": "HISTORY_OF_PRESENTING_COMPLAINT",
            "progress": "MENTAL_STATE",
            "risk_harm": "RISK",
            "risk_property": "RISK",
            "forensic": "FORENSIC",
        }

        # Alternative names for categories (comprehensive list)
        CATEGORY_ALIASES = {
            # Past Psychiatric History variations
            "Psychiatric History": "PAST_PSYCH",
            "Past Psychiatric History": "PAST_PSYCH",
            "Background": "PAST_PSYCH",
            "Background History": "PAST_PSYCH",
            "Previous Psychiatric History": "PAST_PSYCH",
            "Mental Health History": "PAST_PSYCH",
            "Past Mental Health History": "PAST_PSYCH",
            "PAST_PSYCH": "PAST_PSYCH",

            # History of Presenting Complaint variations
            "History of Presenting Complaint": "HISTORY_OF_PRESENTING_COMPLAINT",
            "HPC": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Presenting Complaint": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Current Presentation": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Circumstances of Admission": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Current Admission": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Admission": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Reason for Admission": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Presenting Problem": "HISTORY_OF_PRESENTING_COMPLAINT",
            "Current Episode": "HISTORY_OF_PRESENTING_COMPLAINT",
            "History": "HISTORY_OF_PRESENTING_COMPLAINT",

            # Mental State variations
            "Mental State": "MENTAL_STATE",
            "Mental State Examination": "MENTAL_STATE",
            "MSE": "MENTAL_STATE",
            "Current Mental State": "MENTAL_STATE",

            # Forensic variations
            "Forensic History": "FORENSIC",
            "Forensic": "FORENSIC",

            # Risk variations
            "Risk": "RISK",
            "Risk Assessment": "RISK",

            # Summary
            "Summary": "SUMMARY",
            "Summary and Plan": "SUMMARY",
            "Plan": "SUMMARY",
        }

        # Print all available categories for debugging
        print(f"[TRIBUNAL] Available category aliases mapping:")
        for alias, canonical in CATEGORY_ALIASES.items():
            if canonical == SECTION_TO_CATEGORY.get(key):
                print(f"[TRIBUNAL]   ✓ '{alias}' -> '{canonical}' (MATCHES current section)")

        def extract_text_from_items(category_data):
            """Extract text from category data (handles multiple formats)."""
            parts = []
            if isinstance(category_data, dict) and "items" in category_data:
                for item in category_data.get("items", []):
                    if isinstance(item, dict) and item.get("text"):
                        parts.append(item["text"])
                    elif isinstance(item, str):
                        parts.append(item)
            elif isinstance(category_data, list):
                for item in category_data:
                    if isinstance(item, dict) and item.get("text"):
                        parts.append(item["text"])
                    elif isinstance(item, str):
                        parts.append(item)
            elif isinstance(category_data, str):
                parts.append(category_data)
            return "\n\n".join(parts)

        def find_category_data(target_category):
            """Find data for a category, checking aliases (case-insensitive)."""
            print(f"[TRIBUNAL] find_category_data looking for: '{target_category}'")
            print(f"[TRIBUNAL]   Categories in data: {list(categories.keys())}")

            # Direct match (case-insensitive)
            for cat_key in categories.keys():
                if cat_key.lower() == target_category.lower():
                    print(f"[TRIBUNAL]   ✓ Direct match found: '{cat_key}'")
                    return categories[cat_key]

            # Check aliases (case-insensitive exact match)
            for alias, canonical in CATEGORY_ALIASES.items():
                if canonical.lower() == target_category.lower():
                    for cat_key in categories.keys():
                        if cat_key.lower() == alias.lower():
                            print(f"[TRIBUNAL]   ✓ Alias match: '{cat_key}' -> '{canonical}'")
                            return categories[cat_key]

            # Partial match - check if alias is contained in category key or vice versa
            for alias, canonical in CATEGORY_ALIASES.items():
                if canonical.lower() == target_category.lower():
                    for cat_key in categories.keys():
                        if alias.lower() in cat_key.lower() or cat_key.lower() in alias.lower():
                            print(f"[TRIBUNAL]   ✓ Partial match: '{cat_key}' contains '{alias}'")
                            return categories[cat_key]

            print(f"[TRIBUNAL]   ✗ No match found for '{target_category}'")
            return None

        # Store extracted data globally for ALL sections to access
        self._extracted_categories = categories
        print(f"[TRIBUNAL] *** STORED extracted data with {len(categories)} categories ***")
        print(f"[TRIBUNAL] *** Category keys stored: {list(categories.keys())} ***")

        # ONLY fill the CURRENT section's card (not all cards)
        sections_filled = []

        if key in SECTION_TO_CATEGORY:
            primary_category = SECTION_TO_CATEGORY[key]
            print(f"[TRIBUNAL] Filling ONLY current section '{key}' with '{primary_category}'")

            category_data = find_category_data(primary_category)
            if category_data:
                text = extract_text_from_items(category_data)
                if text.strip():
                    self.cards[key].editor.setPlainText(text)
                    sections_filled.append(key)
                    print(f"[TRIBUNAL] ✓ Filled section '{key}' with {primary_category} data")

        print(f"[TRIBUNAL] Sections filled: {sections_filled}")
        print(f"[TRIBUNAL] Other sections can access data via their popups")
        print(f"{'='*60}\n")

        # Advance to next section
        if sections_filled:
            self._advance_to_next_section(key)

    def _filter_entries_around_admission(self, categories: dict, days_before: int = 5, days_after: int = 3) -> dict:
        """
        Filter entries to show only those around the current admission date.
        Finds the most recent admission date and returns entries within the date range.

        Args:
            categories: Dict of categories with items
            days_before: Number of days before admission to include
            days_after: Number of days after admission to include

        Returns:
            Filtered categories dict
        """
        print(f"[TRIBUNAL] Filtering entries around admission date...")

        def parse_date(date_val):
            """Parse various date formats."""
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            if isinstance(date_val, str):
                # Try common formats
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%dT%H:%M:%S", "%d/%m/%y"]:
                    try:
                        return datetime.strptime(date_val.split()[0], fmt)
                    except:
                        continue
            return None

        # Step 1: Find the most recent admission date from all entries
        # Look for keywords like "admitted", "admission", and take the most recent dated entry
        all_dates = []
        admission_dates = []

        for cat_name, cat_data in categories.items():
            if not isinstance(cat_data, dict):
                continue
            for item in cat_data.get("items", []):
                item_date = parse_date(item.get("date"))
                if item_date:
                    all_dates.append(item_date)
                    # Check if this entry mentions admission
                    text = (item.get("text") or "").lower()
                    if any(kw in text for kw in ["admitted", "admission", "sectioned", "detained", "mha"]):
                        admission_dates.append(item_date)

        if not all_dates:
            print(f"[TRIBUNAL] No dated entries found - returning all data")
            return categories

        # Use the most recent admission date, or if none found, the most recent date overall
        if admission_dates:
            current_admission_date = max(admission_dates)
            print(f"[TRIBUNAL] Found admission date from keywords: {current_admission_date.strftime('%d/%m/%Y')}")
        else:
            # Fall back to most recent date in the data
            current_admission_date = max(all_dates)
            print(f"[TRIBUNAL] Using most recent date as admission: {current_admission_date.strftime('%d/%m/%Y')}")

        # Step 2: Define the date window
        window_start = current_admission_date - timedelta(days=days_before)
        window_end = current_admission_date + timedelta(days=days_after)
        print(f"[TRIBUNAL] Date window: {window_start.strftime('%d/%m/%Y')} to {window_end.strftime('%d/%m/%Y')}")

        # Step 3: Filter entries within the window
        filtered_categories = {}
        total_before = 0
        total_after = 0

        for cat_name, cat_data in categories.items():
            if not isinstance(cat_data, dict):
                continue

            filtered_items = []
            original_items = cat_data.get("items", [])
            total_before += len(original_items)

            for item in original_items:
                item_date = parse_date(item.get("date"))
                if item_date:
                    if window_start <= item_date <= window_end:
                        filtered_items.append(item)
                else:
                    # Include undated items that mention current circumstances
                    text = (item.get("text") or "").lower()
                    if any(kw in text for kw in ["current", "recent", "this admission", "currently"]):
                        filtered_items.append(item)

            total_after += len(filtered_items)

            if filtered_items:
                filtered_categories[cat_name] = {
                    "name": cat_data.get("name", cat_name),
                    "items": filtered_items
                }

        print(f"[TRIBUNAL] Filtered: {total_before} entries -> {total_after} entries around admission")
        return filtered_categories

    def _filter_entries_around_admission_with_info(self, categories: dict, days_before: int = 5, days_after: int = 3) -> tuple:
        """
        Filter entries around admission date and return date info for display.

        Returns:
            Tuple of (filtered_categories dict, date_info string)
        """
        print(f"[TRIBUNAL] Filtering entries around admission date (with info)...")

        def parse_date(date_val):
            """Parse various date formats."""
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            if isinstance(date_val, str):
                for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%dT%H:%M:%S", "%d/%m/%y"]:
                    try:
                        return datetime.strptime(date_val.split()[0], fmt)
                    except:
                        continue
            return None

        # Find admission date
        all_dates = []
        admission_dates = []

        for cat_name, cat_data in categories.items():
            if not isinstance(cat_data, dict):
                continue
            for item in cat_data.get("items", []):
                item_date = parse_date(item.get("date"))
                if item_date:
                    all_dates.append(item_date)
                    text = (item.get("text") or "").lower()
                    if any(kw in text for kw in ["admitted", "admission", "sectioned", "detained", "mha"]):
                        admission_dates.append(item_date)

        if not all_dates:
            return categories, "No dated entries found"

        if admission_dates:
            current_admission_date = max(admission_dates)
        else:
            current_admission_date = max(all_dates)

        window_start = current_admission_date - timedelta(days=days_before)
        window_end = current_admission_date + timedelta(days=days_after)

        date_info = f"Admission: {current_admission_date.strftime('%d/%m/%Y')} | Showing: {window_start.strftime('%d/%m/%Y')} to {window_end.strftime('%d/%m/%Y')}"

        # Filter entries
        filtered_categories = {}
        total_entries = 0

        for cat_name, cat_data in categories.items():
            if not isinstance(cat_data, dict):
                continue

            filtered_items = []
            for item in cat_data.get("items", []):
                item_date = parse_date(item.get("date"))
                if item_date and window_start <= item_date <= window_end:
                    filtered_items.append(item)
                elif not item_date:
                    text = (item.get("text") or "").lower()
                    if any(kw in text for kw in ["current", "recent", "this admission", "currently"]):
                        filtered_items.append(item)

            total_entries += len(filtered_items)

            if filtered_items:
                filtered_categories[cat_name] = {
                    "name": cat_data.get("name", cat_name),
                    "items": filtered_items
                }

        date_info += f" | {total_entries} entries"
        print(f"[TRIBUNAL] {date_info}")
        return filtered_categories, date_info

    def _filter_raw_notes_around_admission(self, notes: list, days_before: int = 5, days_after: int = 3) -> tuple:
        """
        Filter raw notes to show only those around the current admission date.
        Uses build_timeline to accurately identify admission dates.

        Args:
            notes: List of note dicts with 'date' and 'text'/'content' fields
            days_before: Days before admission to include
            days_after: Days after admission to include

        Returns:
            Tuple of (filtered_notes list, date_info string)
        """
        from timeline_builder import build_timeline
        from datetime import date as date_type

        print(f"[TRIBUNAL] Filtering {len(notes)} raw notes around admission date...")

        def parse_date(date_val):
            """Parse various date formats to datetime."""
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            if isinstance(date_val, date_type) and not isinstance(date_val, datetime):
                return datetime.combine(date_val, datetime.min.time())
            if isinstance(date_val, str):
                # Try multiple date formats
                for fmt in ["%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y",
                            "%Y-%m-%dT%H:%M:%S", "%d/%m/%y", "%Y-%m-%d %H:%M:%S.%f"]:
                    try:
                        return datetime.strptime(str(date_val).strip(), fmt)
                    except:
                        continue
                # Try just the date part
                try:
                    return datetime.strptime(str(date_val).split()[0], "%Y-%m-%d")
                except:
                    pass
            # Try parsing datetime objects that came as other types
            try:
                if hasattr(date_val, 'date'):
                    return datetime.combine(date_val.date(), datetime.min.time())
            except:
                pass
            return None

        # Step 1: Prepare notes with proper datetime objects (same as data extractor)
        prepared_notes = []
        for note in notes:
            note_date = parse_date(note.get("date"))
            if note_date:
                prepared_note = dict(note)
                prepared_note["date"] = note_date
                prepared_notes.append(prepared_note)
            else:
                prepared_notes.append(note)

        print(f"[TRIBUNAL] Prepared {len(prepared_notes)} notes with parsed dates")

        # Step 2: Use build_timeline to find inpatient episodes (same as data extractor)
        inpatient_episodes = []
        try:
            episodes = build_timeline(prepared_notes)
            print(f"[TRIBUNAL] build_timeline returned {len(episodes)} episodes")

            for ep in episodes:
                ep_type = ep.get("type", "")
                start = ep.get("start")
                end = ep.get("end")
                print(f"[TRIBUNAL]   Episode: type={ep_type}, start={start}")

                if ep_type == "inpatient" and start:
                    # Convert dates to datetime if needed
                    if isinstance(start, date_type) and not isinstance(start, datetime):
                        start = datetime.combine(start, datetime.min.time())
                    if isinstance(end, date_type) and not isinstance(end, datetime):
                        end = datetime.combine(end, datetime.min.time())
                    inpatient_episodes.append({"start": start, "end": end})
                    print(f"[TRIBUNAL]   -> Inpatient admission: {start} to {end}")

        except Exception as e:
            print(f"[TRIBUNAL] Timeline builder error: {e}")
            import traceback
            traceback.print_exc()

        if not inpatient_episodes:
            print(f"[TRIBUNAL] No inpatient admissions found from timeline")
            # Fallback: use the most recent note date
            all_dates = []
            for note in prepared_notes:
                note_date = note.get("date")
                if isinstance(note_date, datetime):
                    all_dates.append(note_date)
            if all_dates:
                current_admission_date = max(all_dates)
                window_end = current_admission_date + timedelta(days=days_after)
                print(f"[TRIBUNAL] Fallback: using most recent note date: {current_admission_date.strftime('%d/%m/%Y')}")
            else:
                return notes[:50], f"No dates found - showing first 50 notes"
        else:
            # Get the MOST RECENT (LAST) admission by start date
            most_recent = max(inpatient_episodes, key=lambda ep: ep["start"])
            current_admission_date = most_recent["start"]
            # Use actual admission end date from timeline, not a fixed offset
            window_end = most_recent["end"] if most_recent["end"] else current_admission_date + timedelta(days=days_after)
            print(f"[TRIBUNAL] Most recent admission: {current_admission_date.strftime('%d/%m/%Y')} to {window_end.strftime('%d/%m/%Y')}")

        # Step 3: Define date window
        window_start = current_admission_date - timedelta(days=days_before)
        print(f"[TRIBUNAL] Window: {window_start.strftime('%d/%m/%Y')} to {window_end.strftime('%d/%m/%Y')}")

        # Step 4: Filter notes within the window
        filtered_notes = []
        for note in prepared_notes:
            note_date = note.get("date")
            if isinstance(note_date, datetime) and window_start <= note_date <= window_end:
                filtered_notes.append(note)
                print(f"[TRIBUNAL]   Including note from {note_date.strftime('%d/%m/%Y')}: {str(note.get('text', ''))[:50]}...")

        print(f"[TRIBUNAL] Filtered to {len(filtered_notes)} notes in window")

        # Sort by date (oldest first for chronological reading)
        filtered_notes.sort(key=lambda n: n.get("date") or datetime.min)

        date_info = f"Admission: {current_admission_date.strftime('%d/%m/%Y')} | {window_start.strftime('%d/%m/%Y')} to {window_end.strftime('%d/%m/%Y')} | {len(filtered_notes)} notes"
        print(f"[TRIBUNAL] {date_info}")

        return filtered_notes, date_info

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            from shared_data_store import get_shared_store
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file - auto-detect PDF tribunal forms vs other files."""
        # Check if it's a PDF - try XFA extraction first
        if file_path.lower().endswith('.pdf'):
            try:
                from pdf_loader import load_tribunal_pdf, format_radio_value

                # Try to extract XFA form data
                result = load_tribunal_pdf(file_path)

                # If we got T131 data, populate the sections
                if result.get('form_type') == 'T131' and result.get('sections'):
                    self._populate_from_pdf(result, file_path)
                    return
                # If it's a PDF but not a T131 form, send to data extractor
                else:
                    print(f"[TRIBUNAL] PDF is not a T131 form, sending to Data Extractor")
                    self._send_to_data_extractor(file_path)
                    return
            except ImportError:
                # pdf_loader not available, fall back to data extractor
                print("[TRIBUNAL] pdf_loader not available, using Data Extractor")
                self._send_to_data_extractor(file_path)
                return
            except Exception as e:
                print(f"[TRIBUNAL] PDF extraction failed: {e}, using Data Extractor")
                self._send_to_data_extractor(file_path)
                return
        elif file_path.lower().endswith(('.docx', '.doc')):
            # Try to parse DOCX tribunal report
            try:
                result = self._parse_tribunal_docx(file_path)
                if result.get('sections'):
                    self._populate_from_docx(result, file_path)
                    return
                else:
                    print("[TRIBUNAL] No sections found in DOCX, using Data Extractor")
                    self._send_to_data_extractor(file_path)
                    return
            except Exception as e:
                print(f"[TRIBUNAL] DOCX parsing failed: {e}, using Data Extractor")
                self._send_to_data_extractor(file_path)
                return
        else:
            # Other files go to data extractor
            self._send_to_data_extractor(file_path)

    def _send_to_data_extractor(self, file_path):
        """Send a file to the data extractor for processing."""
        self._data_extractor_source_file = file_path
        # Open the data extractor overlay
        self._open_data_extractor_overlay()

        # Pass the file to the data extractor
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            if hasattr(self._data_extractor_overlay, 'load_file'):
                self._data_extractor_overlay.load_file(file_path)
            elif hasattr(self._data_extractor_overlay, 'upload_and_extract'):
                # Fallback - trigger upload dialog (user will need to select again)
                self._data_extractor_overlay.upload_and_extract()
            print(f"[TRIBUNAL] Sent to Data Extractor: {file_path}")

    def _ask_import_action(self, source_filename: str, import_type: str = "report") -> str:
        """Ask user whether to add to existing imported data or replace it.

        Returns: 'add', 'replace', or 'cancel'
        """
        from PySide6.QtWidgets import QMessageBox

        if import_type == "report":
            has_existing = hasattr(self, '_imported_report_data') and any(self._imported_report_data.values())
        else:
            has_existing = hasattr(self, '_extracted_raw_notes') and bool(self._extracted_raw_notes)

        if not has_existing:
            return "replace"

        msg = QMessageBox(self)
        msg.setWindowTitle("Import Data")
        if import_type == "report":
            msg.setText(
                f"Report data has already been imported.\n\n"
                f"New file: {source_filename}\n\n"
                f"Would you like to add this data to the existing import, or replace it?"
            )
        else:
            msg.setText(
                f"Clinical notes have already been loaded.\n\n"
                f"Would you like to add these notes to the existing set, or replace them?"
            )
        add_btn = msg.addButton("Add to Existing", QMessageBox.AcceptRole)
        replace_btn = msg.addButton("Replace All", QMessageBox.DestructiveRole)
        cancel_btn = msg.addButton("Cancel", QMessageBox.RejectRole)
        msg.setDefaultButton(add_btn)
        msg.exec()

        clicked = msg.clickedButton()
        if clicked == cancel_btn:
            return "cancel"
        elif clicked == replace_btn:
            return "replace"
        return "add"

    def _merge_report_section(self, key: str, new_content: str, source_filename: str) -> str:
        """Merge new report content with existing section, adding source reference."""
        existing = self._imported_report_data.get(key, "")
        if not existing:
            return new_content
        return f"{existing}\n\n─── Source: {source_filename} ───\n\n{new_content}"

    def _populate_from_pdf(self, result, file_path):
        """Populate report sections from PDF extraction result."""
        from PySide6.QtWidgets import QMessageBox
        from pdf_loader import format_radio_value
        import os

        sections = result.get('sections', {})

        # Map section keys to card keys
        section_to_card = {
            'patient_details': 'patient_details',
            'author': 'author',
            'factors_hearing': 'factors_hearing',
            'adjustments': 'adjustments',
            'forensic': 'forensic',
            'previous_mh_dates': 'previous_mh_dates',
            'previous_admission_reasons': 'previous_admission_reasons',
            'current_admission': 'current_admission',
            'diagnosis': 'diagnosis',
            'learning_disability': 'learning_disability',
            'detention_required': 'detention_required',
            'treatment': 'treatment',
            'strengths': 'strengths',
            'progress': 'progress',
            'compliance': 'compliance',
            'mca_dol': 'mca_dol',
            'risk_harm': 'risk_harm',
            'risk_property': 'risk_property',
            's2_detention': 's2_detention',
            'other_detention': 'other_detention',
            'discharge_risk': 'discharge_risk',
            'community': 'community',
            'recommendations': 'recommendations',
            'signature_date': 'signature',
        }

        # Store imported data (cards are NOT auto-filled; user sends from popups)
        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        filename = os.path.basename(file_path)
        action = self._ask_import_action(filename, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        mapped_sections = {}
        for section_key, content in sections.items():
            card_key = section_to_card.get(section_key)
            if card_key:
                # Format radio button values before storing
                if section_key in ['learning_disability', 'detention_required', 's2_detention', 'other_detention']:
                    content = format_radio_value(content)
                if action == "add":
                    content = self._merge_report_section(card_key, content, filename)
                self._imported_report_data[card_key] = content
                mapped_sections[card_key] = content
                print(f"[PDF] Stored T131 section '{card_key}' (ready in popup)")

        # Populate popups with imported data so user can review and send
        self._populate_popups_from_import(mapped_sections)

        # Push patient details to shared store for other forms
        self._push_patient_details_to_shared_store(sections)

        # Show success message
        mapped_count = len(mapped_sections)
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "T131 Report Loaded",
            f"Successfully {action_word.lower()} T131 report from:\n{filename}\n\n"
            f"{action_word} {mapped_count} sections to popups.\n\n"
            f"Click each card to review and send the content."
        )

        print(f"[PDF] Mapped {mapped_count} sections from T131 report to popups")

    def _push_patient_details_to_shared_store(self, sections: dict):
        """Extract patient details from loaded sections and push to shared store."""
        import re
        from shared_data_store import get_shared_store

        patient_details = sections.get('patient_details', '')
        if not patient_details:
            return

        patient_info = {
            "name": None,
            "dob": None,
            "nhs_number": None,
            "gender": None,
            "age": None,
            "ethnicity": None,
            "address": None,
            "mha_status": None,
        }

        from datetime import datetime

        # Extract name
        name_match = re.search(r"(?:Name|Patient)[:\s]+([A-Za-z][A-Za-z\-\' ]+)", patient_details, re.IGNORECASE)
        if name_match:
            patient_info["name"] = name_match.group(1).strip()

        # Extract DOB - try numeric format first
        dob_match = re.search(r"(?:DOB|Date of Birth|D\.O\.B)[:\s]+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})", patient_details, re.IGNORECASE)
        if dob_match:
            dob_str = dob_match.group(1).strip()
            for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y"]:
                try:
                    patient_info["dob"] = datetime.strptime(dob_str, fmt)
                    break
                except ValueError:
                    continue

        # Try text date format: "7 October 1979", "15 Jan 1985", etc.
        if not patient_info["dob"]:
            text_dob_match = re.search(
                r"(?:DOB|Date of Birth|D\.O\.B)[:\s]+(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{4})",
                patient_details, re.IGNORECASE
            )
            if text_dob_match:
                day = int(text_dob_match.group(1))
                month_str = text_dob_match.group(2)
                year = int(text_dob_match.group(3))
                month_map = {
                    'january': 1, 'jan': 1, 'february': 2, 'feb': 2, 'march': 3, 'mar': 3,
                    'april': 4, 'apr': 4, 'may': 5, 'june': 6, 'jun': 6,
                    'july': 7, 'jul': 7, 'august': 8, 'aug': 8, 'september': 9, 'sep': 9,
                    'october': 10, 'oct': 10, 'november': 11, 'nov': 11, 'december': 12, 'dec': 12
                }
                month = month_map.get(month_str.lower(), 1)
                try:
                    patient_info["dob"] = datetime(year, month, day)
                except ValueError:
                    pass

        # Extract address
        address_match = re.search(r"(?:Address|Residence|Place of Residence)[:\s]+(.+?)(?:\n|MHA|NHS|$)", patient_details, re.IGNORECASE)
        if address_match:
            patient_info["address"] = address_match.group(1).strip()

        # Extract MHA status
        mha_match = re.search(r"(?:MHA Status|Mental Health Act|Section)[:\s]+([^\n]+)", patient_details, re.IGNORECASE)
        if mha_match:
            patient_info["mha_status"] = mha_match.group(1).strip()

        # Extract NHS number
        nhs_match = re.search(r"(?:NHS)[:\s]*(\d{3}\s*\d{3}\s*\d{4}|\d{10})", patient_details, re.IGNORECASE)
        if nhs_match:
            nhs = nhs_match.group(1).replace(" ", "")
            patient_info["nhs_number"] = f"{nhs[:3]} {nhs[3:6]} {nhs[6:]}" if len(nhs) == 10 else nhs

        # Extract gender
        gender_match = re.search(r"(?:Gender|Sex)[:\s]*(Male|Female|M|F)\b", patient_details, re.IGNORECASE)
        if gender_match:
            g = gender_match.group(1).upper()
            patient_info["gender"] = "Male" if g in ("MALE", "M") else "Female" if g in ("FEMALE", "F") else None

        # Extract age
        age_patterns = [
            r"(?:AGE)[:\s]*(\d{1,3})\s*(?:years?|yrs?|y\.?o\.?)?\b",
            r"\b(\d{1,3})\s*(?:year|yr)\s*old\b",
            r"\b(\d{1,3})\s*y\.?o\.?\b",
            r"\baged?\s*(\d{1,3})\b",
        ]
        for pattern in age_patterns:
            match = re.search(pattern, patient_details, re.IGNORECASE)
            if match:
                age_val = int(match.group(1))
                if 0 < age_val < 120:
                    patient_info["age"] = age_val
                    break

        # Calculate age from DOB if not found explicitly
        if not patient_info["age"] and patient_info["dob"]:
            today = datetime.today()
            dob = patient_info["dob"]
            age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if 0 < age < 120:
                patient_info["age"] = age

        # Extract ethnicity
        ethnicity_patterns = [
            r"(?:ETHNICITY|ETHNIC\s*(?:GROUP|ORIGIN)?)[:\s]+([A-Za-z][A-Za-z\s\-\/]+?)(?:\n|$|,)",
            r"\b(White\s*(?:British|Irish|European|Other)?)\b",
            r"\b(Black\s*(?:British|African|Caribbean|Other)?)\b",
            r"\b(Asian\s*(?:British|Indian|Pakistani|Bangladeshi|Chinese|Other)?)\b",
            r"\b(Mixed\s*(?:White\s*(?:and|&)\s*(?:Black\s*(?:Caribbean|African)|Asian))?)\b",
        ]
        for pattern in ethnicity_patterns:
            match = re.search(pattern, patient_details, re.IGNORECASE)
            if match:
                ethnicity_val = match.group(1).strip()
                if len(ethnicity_val) > 2:
                    patient_info["ethnicity"] = ethnicity_val.title()
                    break

        # Push to shared store if any info found
        if any(patient_info.values()):
            shared_store = get_shared_store()
            shared_store.set_patient_info(patient_info, source="tribunal_report")
            print(f"[TRIBUNAL] Pushed patient details to SharedDataStore: {list(k for k,v in patient_info.items() if v)}")

    def _get_heading_to_card_mapping(self) -> list:
        """Get heading patterns mapped to T131 card keys."""
        return [
            # Factors affecting hearing
            (r'factors.*affect.*understanding', 'factors_hearing'),
            (r'ability.*cope.*hearing', 'factors_hearing'),

            # Adjustments
            (r'adjustments.*panel.*consider', 'adjustments'),
            (r'adjustments.*may.*consider', 'adjustments'),

            # Index offence/forensic history
            (r'index offence', 'forensic'),
            (r'forensic.*history', 'forensic'),
            (r'relevant.*forensic', 'forensic'),

            # Previous mental health dates
            (r'dates.*previous.*mental health', 'previous_mh_dates'),
            (r'previous.*involvement.*mental health', 'previous_mh_dates'),

            # Previous admission reasons
            (r'reasons.*previous.*admission', 'previous_admission_reasons'),
            (r'give reasons.*previous', 'previous_admission_reasons'),

            # Current admission
            (r'circumstances.*current.*admission', 'current_admission'),
            (r'current admission', 'current_admission'),

            # Diagnosis
            (r'is the patient.*suffering', 'diagnosis'),
            (r'mental disorder.*nature', 'diagnosis'),
            (r'9\.\s*mental disorder', 'diagnosis'),
            (r'9\.\s*is the patient', 'diagnosis'),
            (r'nature.*degree.*mental disorder', 'diagnosis'),
            (r'from.*mental disorder', 'diagnosis'),
            (r'diagnosis', 'diagnosis'),

            # Learning disability
            (r'learning disability', 'learning_disability'),
            (r'abnormally aggressive', 'learning_disability'),

            # Detention requirement
            (r'what is it.*necessary', 'detention_required'),
            (r'necessary.*medical treatment', 'detention_required'),
            (r'why is detention necessary', 'detention_required'),

            # Treatment
            (r'medical treatment.*prescribed', 'treatment'),
            (r'appropriate.*available.*treatment', 'treatment'),
            (r'what.*treatment', 'treatment'),

            # Strengths
            (r'strengths.*positive factors', 'strengths'),
            (r'what are the strengths', 'strengths'),

            # Progress
            (r'current progress', 'progress'),
            (r'summary.*progress', 'progress'),
            (r'progress.*behaviour', 'progress'),

            # Compliance (Section 15)
            (r'understanding.*compliance', 'compliance'),
            (r'compliance.*willingness', 'compliance'),
            (r'willingness to accept', 'compliance'),
            (r"patient's understanding of", 'compliance'),
            (r'future willingness to accept', 'compliance'),
            (r'prescribed medication.*mental disorder', 'compliance'),
            (r'comply with.*medical treatment', 'compliance'),
            (r'15\.\s*what is the patient', 'compliance'),
            (r'what is the patient.*understanding', 'compliance'),

            # MCA/DoL
            (r'mental capacity act', 'mca_dol'),
            (r'deprivation of liberty', 'mca_dol'),
            (r'mca.*dol', 'mca_dol'),

            # Risk harm
            (r'harmed themselves or others', 'risk_harm'),
            (r'incidents.*harm', 'risk_harm'),
            (r'threatened to harm', 'risk_harm'),

            # Risk property
            (r'damaged property', 'risk_property'),
            (r'threatened to damage property', 'risk_property'),

            # Section 2 detention
            (r'section 2 cases.*detention', 's2_detention'),
            (r'in section 2 cases', 's2_detention'),

            # Other detention
            (r'all other cases.*provision', 'other_detention'),
            (r'in all other cases', 'other_detention'),

            # Discharge risk
            (r'discharged.*dangerous', 'discharge_risk'),
            (r'likely to act in a manner dangerous', 'discharge_risk'),

            # Community
            (r'risks.*managed.*community', 'community'),
            (r'managed effectively in the community', 'community'),

            # Recommendations
            (r'recommendations.*tribunal', 'recommendations'),
            (r'do you have any recommendations', 'recommendations'),

            # Other info
            (r'other relevant information', 'other_info'),
            (r'other.*information.*tribunal', 'other_info'),
        ]

    def _match_heading_to_card(self, heading_text: str) -> str:
        """Match a heading to a T131 card key using regex patterns."""
        import re

        if not heading_text or len(heading_text) < 10:
            return None

        lower = heading_text.lower()
        patterns = self._get_heading_to_card_mapping()

        for pattern, card_key in patterns:
            if re.search(pattern, lower):
                return card_key

        return None

    def _parse_tribunal_docx(self, file_path: str) -> dict:
        """Parse a DOCX tribunal report and extract sections."""
        from docx import Document
        import re

        try:
            doc = Document(file_path)
        except Exception as e:
            print(f"[TRIBUNAL] Failed to open DOCX: {e}")
            return {}

        result = {'form_type': 'T131', 'sections': {}}

        def get_unique_cells(cell_list):
            unique = []
            for c in cell_list:
                if c and (not unique or c != unique[-1]):
                    unique.append(c)
            return unique

        def is_question_text(text):
            if not text or len(text) < 10:
                return False
            lower = text.lower().strip()
            # Remove leading checkbox/number characters to check the actual text
            lower_clean = re.sub(r'^[\[\]☐☒xX\s\-–\d\.]*', '', lower).strip()

            # Question patterns - check at START of cleaned text only
            question_patterns = [
                'are there any factors', 'are there any adjustments',
                'what is the nature of', 'give details of any',
                'what are the strengths', 'give a summary of',
                'in section 2 cases', 'in all other cases',
                'if the patient was discharged', 'if the patient were discharged',
                'please explain how', 'is there any other relevant',
                'do you have any recommendations', 'is the patient now suffering',
                'what appropriate and available', 'what are the dates',
                'what are the circumstances', 'give reasons for',
                'does the patient have a learning', 'what is it about',
                'would they be likely to act', 'managed effectively in the community',
                'if yes, has a diagnosis', 'if yes, what is the diagnosis',
                'has a diagnosis been made', 'what is the diagnosis',
            ]
            for pattern in question_patterns:
                if lower_clean.startswith(pattern):
                    return True
            if re.match(r'^\d{1,2}\.\s+', text):
                return True
            return False

        def is_checkbox_only(text):
            """Check if text is just checkbox format without actual content."""
            if not text:
                return False
            # Remove checkbox symbols and whitespace
            cleaned = re.sub(r'[\[\]☐☒xX\s\n]', '', text)
            # Remove common labels
            cleaned = re.sub(r'(?:No|Yes|N/A)', '', cleaned, flags=re.IGNORECASE)
            # If nothing substantial left, it's checkbox only
            return len(cleaned.strip()) < 5

        # Parse tables
        for table in doc.tables:
            rows = list(table.rows)
            i = 0
            while i < len(rows):
                cells = [cell.text.strip() for cell in rows[i].cells]
                if not any(cells):
                    i += 1
                    continue

                unique_cells = get_unique_cells(cells)

                # Try all cells for heading match
                heading_key = None
                heading_cell_idx = -1
                for cell_idx, cell_text in enumerate(unique_cells):
                    if re.match(r'^\d{1,2}\.\s*$', cell_text):
                        continue
                    matched = self._match_heading_to_card(cell_text)
                    if matched:
                        heading_key = matched
                        heading_cell_idx = cell_idx
                        break

                if heading_key:
                    answer = ""

                    # Check cells after heading
                    start_idx = heading_cell_idx + 1 if heading_cell_idx >= 0 else 1
                    if len(unique_cells) > start_idx:
                        for cell_text in unique_cells[start_idx:]:
                            if is_question_text(cell_text):
                                continue
                            if cell_text and not re.match(r'^\d+\.\s*', cell_text):
                                cleaned = self._clean_docx_content(cell_text)
                                if cleaned and cleaned not in ('No', 'Yes', 'N/A'):
                                    answer = cleaned
                                    break

                    # Check next row
                    if not answer and i + 1 < len(rows):
                        next_cells = [cell.text.strip() for cell in rows[i + 1].cells]
                        unique_next = get_unique_cells(next_cells)

                        for cell_text in unique_next:
                            if re.match(r'^\d+\.\s*', cell_text):
                                break
                            if is_question_text(cell_text):
                                continue
                            # Skip Yes/No label cells
                            if cell_text.strip() in ('No\nYes', 'Yes\nNo', 'No', 'Yes', 'N/A'):
                                continue

                            # Check for checkbox format - extract the actual content
                            if '[x' in cell_text.lower() or '☒' in cell_text or '[' in cell_text:
                                yes_answer = self._extract_checkbox_answer(cell_text)
                                if yes_answer:
                                    answer = yes_answer
                                    break

                            cleaned = self._clean_docx_content(cell_text)
                            if cleaned and not is_question_text(cleaned):
                                # Skip if it's just Yes/No labels or checkbox only
                                if cleaned.replace('\n', ' ').strip() in ('No Yes', 'Yes No'):
                                    continue
                                if is_checkbox_only(cleaned):
                                    continue
                                answer = cleaned
                                break

                        if answer:
                            i += 1

                    if answer:
                        # Add Yes prefix for certain sections
                        yes_no_keys = {'factors_hearing', 'adjustments', 's2_detention',
                                      'other_detention', 'discharge_risk', 'recommendations'}
                        if heading_key in yes_no_keys and answer not in ('Yes', 'No', 'N/A'):
                            if not answer.startswith('Yes') and not answer.startswith('No'):
                                answer = f"Yes - {answer}"
                        result['sections'][heading_key] = answer
                        print(f"[TRIBUNAL] Matched heading ({heading_key}): {answer[:50]}...")

                # Check for patient details
                elif 'Name of Patient' in cells[0] or 'name of patient' in cells[0].lower():
                    if len(unique_cells) > 1:
                        result['sections']['patient_details'] = f"Name: {unique_cells[1]}"
                elif 'Date of Birth' in cells[0]:
                    if len(unique_cells) > 1:
                        if 'patient_details' not in result['sections']:
                            result['sections']['patient_details'] = ""
                        result['sections']['patient_details'] += f"\nDOB: {unique_cells[1]}"
                elif 'NHS Number' in cells[0]:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nNHS: {unique_cells[1]}"
                elif 'Usual Place of Residence' in cells[0]:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nAddress: {unique_cells[1]}"
                elif 'Mental Health Act Status' in cells[0]:
                    if len(unique_cells) > 1 and 'patient_details' in result['sections']:
                        result['sections']['patient_details'] += f"\nMHA Status: {unique_cells[1]}"
                elif 'Your Name' in cells[0]:
                    if len(unique_cells) > 1:
                        result['sections']['author'] = f"Name: {unique_cells[1]}"
                elif 'Your Role' in cells[0]:
                    if len(unique_cells) > 1 and 'author' in result['sections']:
                        result['sections']['author'] += f"\nRole: {unique_cells[1]}"
                elif 'Date of Report' in cells[0]:
                    if len(unique_cells) > 1 and 'author' in result['sections']:
                        result['sections']['author'] += f"\nDate: {unique_cells[1]}"

                i += 1

        print(f"[TRIBUNAL] Parsed DOCX: found {len(result['sections'])} sections")
        return result

    def _clean_docx_content(self, text: str) -> str:
        """Clean up content from DOCX."""
        import re
        if not text:
            return ""
        text = re.sub(r'\[\s*[xX]?\s*\]', '', text)
        text = re.sub(r'[☐☒]', '', text)
        text = re.sub(r'–\s*If yes[^:?]*[:\?]?\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _extract_checkbox_answer(self, text: str) -> str:
        """Extract answer from checkbox format like [x] - If yes, explain..."""
        import re
        # Look for checked checkbox followed by content
        match = re.search(r'\[[\s]*[xX][\s]*\][^a-zA-Z]*(?:If yes[^\n]*\n+)?(.+)', text, re.DOTALL | re.IGNORECASE)
        if match:
            content = match.group(1).strip()
            if content and len(content) > 5:
                return content
        # Also check for ☒
        match = re.search(r'☒[^a-zA-Z]*(?:If yes[^\n]*\n+)?(.+)', text, re.DOTALL | re.IGNORECASE)
        if match:
            content = match.group(1).strip()
            if content and len(content) > 5:
                return content
        return ""

    def _populate_from_docx(self, result: dict, file_path: str):
        """Populate report sections from DOCX parsing result."""
        from PySide6.QtWidgets import QMessageBox
        from shared_data_store import get_shared_store
        import os

        sections = result.get('sections', {})

        # Store imported data for popups to use
        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        filename = os.path.basename(file_path)
        action = self._ask_import_action(filename, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        # Store imported data (cards are NOT auto-filled; user sends from popups)
        merged_sections = {}
        for section_key, content in sections.items():
            if action == "add":
                content = self._merge_report_section(section_key, content, filename)
            self._imported_report_data[section_key] = content
            merged_sections[section_key] = content
            print(f"[TRIBUNAL] Stored DOCX section '{section_key}' (ready in popup)")

        # Populate popups with imported data so user can review and send
        self._populate_popups_from_import(merged_sections)

        # Push patient details to shared store
        self._push_patient_details_to_shared_store(sections)

        # Push sections to shared store for cross-talk with nursing form
        shared_store = get_shared_store()
        shared_store.set_report_sections(merged_sections, source_form="tribunal")

        # Show success message
        mapped_count = len(merged_sections)
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"Successfully {action_word.lower()} report from:\n{filename}\n\n"
            f"{action_word} {mapped_count} sections to popups.\n\n"
            f"Click each card to review and send the content."
        )

        print(f"[TRIBUNAL] {action_word} {mapped_count} sections from DOCX to popups")

    def _populate_popups_from_import(self, sections: dict):
        """Populate popups with imported report data."""
        for section_key, content in sections.items():
            if not content:
                continue

            # Create popup if it doesn't exist yet (popups are normally created lazily on click)
            if section_key not in self.popups:
                popup = self._create_popup(section_key)
                if popup:
                    self.popups[section_key] = popup
                    self.popup_stack.addWidget(popup)
                    # Connect signals
                    if hasattr(popup, 'sent'):
                        if section_key == "forensic":
                            popup.sent.connect(lambda text, state, k=section_key: self._update_card(k, text))
                        else:
                            popup.sent.connect(lambda text, k=section_key: self._update_card(k, text))
                    print(f"[TRIBUNAL] Created popup '{section_key}' during import")

            # If popup exists (now guaranteed), populate it
            if section_key in self.popups:
                popup = self.popups[section_key]
                self._populate_single_popup(popup, section_key, content)

    def _populate_single_popup(self, popup, section_key: str, content: str):
        """Populate a single popup with imported content."""
        import re

        # ============================================================
        # HELPER FUNCTIONS FOR CHECKBOX DETECTION
        # ============================================================
        # Checked box symbols: ☒ (U+2612), ☑ (U+2611), ✓ (U+2713), ✔ (U+2714), ✗ (U+2717), ✘ (U+2718)
        # Unchecked box symbols: ☐ (U+2610), □ (U+25A1)
        CHECKED_SYMBOLS = r'[☒☑✓✔✗✘xX]'
        UNCHECKED_SYMBOLS = r'[☐□]'

        def has_yes_cross(text):
            """Check if Yes has a checked box symbol."""
            # Check for [x] or [X] pattern after Yes
            if re.search(r'yes\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            # Check for checked symbol after Yes (with flexible whitespace)
            if re.search(r'yes\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            # Check for checked symbol before Yes
            if re.search(r'[☒☑✓✔]\s*yes', text, re.IGNORECASE):
                return True
            # Check for pattern: Yes with checked box, No with unchecked box
            # e.g., "Yes ☒  No ☐" or "Yes [X]  No [ ]"
            if re.search(r'yes\s*[☒☑✓✔\[xX\]].*no\s*[☐□\[\s\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes.*[☐□]\s*no', text, re.IGNORECASE):
                return True
            return False

        def has_no_cross(text):
            """Check if No has a checked box symbol (meaning answer is No)."""
            # Check for [x] or [X] pattern after No
            if re.search(r'no\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            # Check for checked symbol after No
            if re.search(r'no\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            # Check for checked symbol before No
            if re.search(r'[☒☑✓✔]\s*no\b', text, re.IGNORECASE):
                return True
            # Check for pattern: Yes with unchecked box, No with checked box
            if re.search(r'yes\s*[☐□\[\s\]].*no\s*[☒☑✓✔\[xX\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☐□]\s*yes.*[☒☑✓✔]\s*no', text, re.IGNORECASE):
                return True
            return False

        def is_yes_content(text):
            """Check if content indicates Yes answer."""
            # First check for explicit crosses
            if has_yes_cross(text):
                return True
            if has_no_cross(text):
                return False
            # Fall back to text-based detection
            lower = text.lower().strip()
            if lower.startswith("yes") or "yes -" in lower or "yes," in lower:
                return True
            # If substantial text and not explicitly "no", assume yes
            if len(text.strip()) > 30 and not lower.startswith("no"):
                return True
            return False

        def extract_detail(text):
            """Extract detail text after Yes/No prefix."""
            for prefix in ["Yes - ", "Yes, ", "Yes-", "yes - ", "yes, ", "yes-", "Yes ", "yes "]:
                if text.startswith(prefix):
                    return text[len(prefix):].strip()
            return text

        # ============================================================
        # PATIENT DETAILS - parse text into fields directly
        # ============================================================
        if section_key == "patient_details":
            from datetime import datetime
            patient_info = {}

            name_match = re.search(r"(?:Full\s*Name|Name|Patient)[:\s]+(.+)", content, re.IGNORECASE)
            if name_match:
                patient_info["name"] = name_match.group(1).strip()

            dob_match = re.search(r"(?:Date of Birth|DOB|D\.O\.B)[:\s]+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})", content, re.IGNORECASE)
            if dob_match:
                dob_str = dob_match.group(1).strip()
                for fmt in ["%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y", "%d/%m/%y", "%d-%m-%y"]:
                    try:
                        patient_info["dob"] = datetime.strptime(dob_str, fmt)
                        break
                    except ValueError:
                        continue

            if not patient_info.get("dob"):
                text_dob = re.search(
                    r"(?:Date of Birth|DOB|D\.O\.B)[:\s]+(\d{1,2})\s+(January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+(\d{4})",
                    content, re.IGNORECASE
                )
                if text_dob:
                    month_map = {'january':1,'jan':1,'february':2,'feb':2,'march':3,'mar':3,'april':4,'apr':4,'may':5,'june':6,'jun':6,'july':7,'jul':7,'august':8,'aug':8,'september':9,'sep':9,'october':10,'oct':10,'november':11,'nov':11,'december':12,'dec':12}
                    try:
                        patient_info["dob"] = datetime(int(text_dob.group(3)), month_map.get(text_dob.group(2).lower(), 1), int(text_dob.group(1)))
                    except ValueError:
                        pass

            gender_match = re.search(r"(?:Gender|Sex)[:\s]*(Male|Female|M|F)\b", content, re.IGNORECASE)
            if gender_match:
                g = gender_match.group(1).upper()
                patient_info["gender"] = "Male" if g in ("MALE", "M") else "Female"

            address_match = re.search(r"(?:Usual Place of Residence|Address|Residence)[:\s]+(.+?)(?:\n|$)", content, re.IGNORECASE)
            if address_match:
                patient_info["address"] = address_match.group(1).strip()

            if patient_info and hasattr(popup, 'fill_patient_info'):
                popup.fill_patient_info(patient_info)
                print(f"[TRIBUNAL] Filled patient_details fields: {list(patient_info.keys())}")
            return

        # ============================================================
        # YES/NO POPUPS WITH STANDARD yes_btn/no_btn
        # ============================================================
        yes_no_sections = ["factors_hearing", "adjustments", "diagnosis",
                          "discharge_risk", "recommendations"]

        if section_key in yes_no_sections:
            is_yes = is_yes_content(content)
            if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
                if is_yes:
                    popup.yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set '{section_key}' popup to Yes")
                else:
                    popup.no_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set '{section_key}' popup to No")

                # For factors_hearing, also detect and set the specific factor radio buttons
                if section_key == "factors_hearing" and is_yes:
                    content_lower = content.lower()
                    if hasattr(popup, 'autism_rb') and ('autism' in content_lower or 'autistic' in content_lower):
                        popup.autism_rb.setChecked(True)
                        print(f"[TRIBUNAL] Set factors_hearing to Autism")
                    elif hasattr(popup, 'ld_rb') and ('learning disability' in content_lower or 'learning difficulties' in content_lower):
                        popup.ld_rb.setChecked(True)
                        print(f"[TRIBUNAL] Set factors_hearing to Learning Disability")
                    elif hasattr(popup, 'patience_rb') and ('irritab' in content_lower or 'frustration' in content_lower or 'patience' in content_lower):
                        popup.patience_rb.setChecked(True)
                        print(f"[TRIBUNAL] Set factors_hearing to Low frustration tolerance")

                # Populate the details_field (inside the container) with imported text
                detail = extract_detail(content)
                if detail and detail.lower() not in ("yes", "no"):
                    if hasattr(popup, 'details_field'):
                        popup.details_field.setPlainText(detail)
                        print(f"[TRIBUNAL] Set '{section_key}' details field")
                    elif hasattr(popup, 'additional_details_field'):
                        popup.additional_details_field.setPlainText(detail)

                # Hide the always-visible additional details (empty duplicate)
                if section_key in ("factors_hearing", "adjustments"):
                    if hasattr(popup, 'always_visible_details'):
                        popup.always_visible_details.hide()

                # Send updated text to card
                if hasattr(popup, '_send_to_card'):
                    popup._send_to_card()
                    print(f"[TRIBUNAL] Sent '{section_key}' popup text to card")

        # ============================================================
        # LEARNING DISABILITY (Section 10) - q1_yes_btn, q1_no_btn, q2_*
        # ============================================================
        if section_key == "learning_disability":
            print(f"[TRIBUNAL] Section 10 content: {repr(content[:200] if len(content) > 200 else content)}")
            if hasattr(popup, 'q1_yes_btn') and hasattr(popup, 'q1_no_btn'):
                yes_found = has_yes_cross(content)
                no_found = has_no_cross(content)
                print(f"[TRIBUNAL] Section 10 detection: yes_cross={yes_found}, no_cross={no_found}")
                # Check for Yes [x] vs No [x] pattern
                if yes_found:
                    popup.q1_yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set learning_disability Q1 to Yes (found Yes [x])")
                    # Check Q2 for aggressive/irresponsible
                    if hasattr(popup, 'q2_yes_btn') and hasattr(popup, 'q2_no_btn'):
                        # Look for second Yes [x] or aggressive keywords
                        parts = content.split('\n')
                        for part in parts[1:]:  # Skip first line
                            if has_yes_cross(part) or "aggressive" in part.lower() or "irresponsible" in part.lower():
                                popup.q2_yes_btn.setChecked(True)
                                print(f"[TRIBUNAL] Set learning_disability Q2 to Yes")
                                break
                elif no_found:
                    popup.q1_no_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set learning_disability Q1 to No (found No [x])")
                elif is_yes_content(content):
                    popup.q1_yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set learning_disability Q1 to Yes (text-based)")
                else:
                    # No crosses found - default to No
                    popup.q1_no_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set learning_disability Q1 to No (no crosses found)")

        # ============================================================
        # DETENTION REQUIRED (Section 11) - SimpleYesNoPopup
        # ============================================================
        if section_key == "detention_required":
            print(f"[TRIBUNAL] Section 11 content: {repr(content[:200] if len(content) > 200 else content)}")
            if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
                yes_found = has_yes_cross(content)
                no_found = has_no_cross(content)
                yes_text = is_yes_content(content)
                print(f"[TRIBUNAL] Section 11 detection: yes_cross={yes_found}, no_cross={no_found}, yes_text={yes_text}")
                if yes_found:
                    popup.yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set detention_required to Yes (found Yes [x])")
                elif no_found:
                    popup.no_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set detention_required to No (found No [x])")
                elif yes_text:
                    popup.yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set detention_required to Yes (text-based)")

        # ============================================================
        # S2 DETENTION (Section 19) - YesNoNAPopup
        # ============================================================
        if section_key == "s2_detention":
            print(f"[TRIBUNAL] Section 19 content: {repr(content[:200] if len(content) > 200 else content)}")
            if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
                yes_found = has_yes_cross(content)
                no_found = has_no_cross(content)
                print(f"[TRIBUNAL] Section 19 detection: yes_cross={yes_found}, no_cross={no_found}")
                if yes_found:
                    popup.yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set s2_detention to Yes (found Yes [x])")
                elif no_found:
                    popup.no_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set s2_detention to No (found No [x])")
                elif "n/a" in content.lower() or "not applicable" in content.lower():
                    if hasattr(popup, 'na_btn'):
                        popup.na_btn.setChecked(True)
                        print(f"[TRIBUNAL] Set s2_detention to N/A")
                elif is_yes_content(content):
                    popup.yes_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set s2_detention to Yes (text-based)")

        # ============================================================
        # IMPORTED DATA SECTION - Add collapsible container for ALL popups
        # ============================================================
        # Always add collapsible imported data section (if content is valid)
        # previous_mh_dates (section 6) now uses the same generic path as section 14
        # This ensures consistent UX across all sections
        self._add_imported_data_to_popup(popup, section_key, content)

        # ============================================================
        # SPECIFIC POPUP HANDLING
        # ============================================================

        # Diagnosis popup (section 9) - try to match ICD-10
        if section_key == "diagnosis":
            self._populate_diagnosis_popup(popup, content)

        # Treatment popup (section 12) - extract medications
        if section_key == "treatment":
            self._populate_treatment_popup(popup, content)

        # MCA/DoL popup (section 16)
        if section_key == "mca_dol":
            self._populate_mca_dol_popup(popup, content)

    def _add_imported_data_to_popup(self, popup, section_key: str, content: str):
        """Add imported data collapsible section to popup that doesn't have one."""
        from PySide6.QtWidgets import QLabel, QVBoxLayout, QWidget, QScrollArea
        from PySide6.QtCore import Qt

        # Skip if empty or pointless content like "No [ ] Yes [ ]" or "☐ No ☐ Yes"
        if not content or not content.strip():
            return
        # Skip pointless checkbox-only content
        import re
        cleaned = content.strip()
        # Pattern 1: "No [ ] Yes [ ]" or "Yes [ ] No [ ]"
        if re.match(r'^(No|Yes)\s*\[\s*\]\s*(No|Yes)\s*\[\s*\]$', cleaned, re.IGNORECASE):
            print(f"[TRIBUNAL] Skipping pointless checkbox import for '{section_key}'")
            return
        # Pattern 2: "☐ No ☐ Yes" or similar with checkbox symbols
        if re.match(r'^[☐☒\s]*(No|Yes)\s*[☐☒\s]*(No|Yes)\s*[☐☒\s]*$', cleaned, re.IGNORECASE):
            print(f"[TRIBUNAL] Skipping pointless checkbox import for '{section_key}'")
            return
        # Pattern 3: Just "Yes" or "No" without additional content
        if cleaned.lower() in ('yes', 'no', 'n/a', 'yes.', 'no.'):
            print(f"[TRIBUNAL] Skipping minimal checkbox import for '{section_key}'")
            return

        # Popups with specialized extracted_section that have their own internal data handling
        # These have their own methods to populate extracted_section - skip adding another
        specialized_sections = {
            'treatment'  # TreatmentPopup has its own imported data section with checkboxes
        }
        # Note: progress, compliance, risk_harm, risk_property, discharge_risk now get imported data sections
        # Sections 10 (learning_disability) and 11 (detention_required) are just yes/no click boxes
        # They don't need imported data collapsibles
        yes_no_only_sections = {'learning_disability', 'detention_required'}

        # Sections that don't need imported data collapsibles
        # Sections 19 (s2_detention) and 20 (other_detention) are simple yes/no - no imported data
        skip_imported_data_sections = {'author', 'patient_details', 'factors_hearing', 'adjustments', 's2_detention', 'other_detention', 'signature'}

        # Special handling for forensic - populate its built-in extracted_section with full text
        if section_key == 'forensic':
            self._populate_forensic_extracted_section(popup, content)
            return

        # current_admission (section 8) - hide notes-only sections, then use generic path
        if section_key == 'current_admission':
            if hasattr(popup, 'narrative_section'):
                popup.narrative_section.setVisible(False)
            if hasattr(popup, 'admissions_section'):
                popup.admissions_section.setVisible(False)

        if section_key in specialized_sections:
            # These popups have their own data population - don't add duplicate section
            print(f"[TRIBUNAL] Skipping imported data section for specialized popup '{section_key}'")
            return

        if section_key in yes_no_only_sections:
            # These are just yes/no click sections - no imported data needed
            print(f"[TRIBUNAL] Skipping imported data for yes/no section '{section_key}'")
            return

        if section_key in skip_imported_data_sections:
            # These sections don't need imported data collapsibles
            print(f"[TRIBUNAL] Skipping imported data for section '{section_key}'")
            return

        # Check if we already added imported data
        if hasattr(popup, '_imported_data_added') and popup._imported_data_added:
            # Update existing content instead
            if hasattr(popup, '_imported_content_label'):
                popup._imported_content_label.setText(content)
            return

        try:
            from background_history_popup import CollapsibleSection

            # FIRST: Determine target layout and get parent widget BEFORE creating CollapsibleSection
            # This ensures proper parenting to prevent the widget from appearing as a separate window
            target_layout = None
            parent_widget = None
            layout_source = None

            if hasattr(popup, 'scroll_layout') and popup.scroll_layout:
                target_layout = popup.scroll_layout
                parent_widget = target_layout.parentWidget()
                layout_source = "scroll_layout"
            elif hasattr(popup, 'main_layout') and popup.main_layout:
                target_layout = popup.main_layout
                parent_widget = target_layout.parentWidget()
                layout_source = "main_layout"
            elif hasattr(popup, 'container_layout') and popup.container_layout:
                target_layout = popup.container_layout
                parent_widget = target_layout.parentWidget()
                layout_source = "container_layout"
            elif hasattr(popup, 'layout'):
                # TribunalPopupBase stores layout as attribute, not method
                layout_attr = popup.layout
                if hasattr(layout_attr, 'insertWidget'):
                    # It's a layout object directly
                    target_layout = layout_attr
                    parent_widget = target_layout.parentWidget()
                    layout_source = "layout (attribute)"
                elif callable(layout_attr):
                    # It's the Qt layout() method
                    target_layout = layout_attr()
                    parent_widget = target_layout.parentWidget() if target_layout else None
                    layout_source = "layout() (method)"

            # Fallback: use popup itself as parent if no layout parent found
            if not parent_widget:
                parent_widget = popup
                print(f"[TRIBUNAL] Warning: No layout parent found for '{section_key}', using popup as parent")

            if not target_layout:
                print(f"[TRIBUNAL] No target layout found for '{section_key}' - cannot embed imported data")
                return

            print(f"[TRIBUNAL] Using {layout_source} for '{section_key}', parent: {parent_widget.__class__.__name__}")

            # Create CollapsibleSection WITH parent to ensure proper embedding
            extracted_section = CollapsibleSection("Imported Data", parent=parent_widget, start_collapsed=False)
            extracted_section.setMinimumHeight(40)  # Reduced by 20%
            extracted_section.set_header_style("""
                QFrame {
                    background: rgba(180, 150, 50, 0.25);
                    border: 1px solid rgba(180, 150, 50, 0.5);
                    border-radius: 6px 6px 0 0;
                }
            """)
            extracted_section.title_label.setStyleSheet("""
                QLabel {
                    font-size: 21px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                }
            """)

            # Create content widget
            extracted_content = QWidget()
            extracted_content.setStyleSheet("""
                QWidget {
                    background: rgba(255, 248, 220, 0.95);
                    border: 1px solid rgba(180, 150, 50, 0.4);
                    border-top: none;
                    border-radius: 0 0 12px 12px;
                }
            """)

            extracted_layout = QVBoxLayout(extracted_content)
            extracted_layout.setContentsMargins(12, 10, 12, 10)
            extracted_layout.setSpacing(4)

            # Single text block with one include checkbox in the header
            from PySide6.QtWidgets import QCheckBox

            # Add include checkbox to the header (right side)
            include_cb = QCheckBox()
            include_cb.setToolTip("Include imported data on card")
            include_cb.setStyleSheet("""
                QCheckBox { background: transparent; border: none; }
                QCheckBox::indicator { width: 20px; height: 20px; }
            """)
            include_cb.toggled.connect(lambda checked, sk=section_key, txt=content: self._on_imported_checkbox_toggled(sk, txt, checked))
            extracted_section.header.layout().addWidget(include_cb)
            popup._imported_include_cb = include_cb
            popup._imported_report_text = content

            # Content label — single block of text
            content_label = QLabel(content)
            content_label.setWordWrap(True)
            content_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignTop)
            content_label.setStyleSheet("font-size: 15px; color: #4a4a4a; background: transparent; border: none;")
            content_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
            extracted_layout.addWidget(content_label)
            popup._imported_content_label = content_label

            extracted_section.set_content(extracted_content)
            extracted_section.setVisible(True)
            print(f"[TRIBUNAL] Collapsible content: '{content[:50]}...' ({len(content)} chars)")

            popup.extracted_section = extracted_section

            # Add to popup layout - always at TOP so imported data is visible first
            target_layout.insertWidget(0, extracted_section)
            # Push content to top (absorb extra space from widgetResizable scroll area)
            target_layout.addStretch()
            print(f"[TRIBUNAL] Added CollapsibleSection at TOP of '{section_key}'")

            # Ensure proper parenting after insertion (safety measure)
            if extracted_section.parent() is None:
                extracted_section.setParent(parent_widget)
                print(f"[TRIBUNAL] Warning: Re-parented CollapsibleSection for '{section_key}'")

            extracted_section.setVisible(True)  # Make sure it's visible
            popup._imported_data_added = True

            # Force layout update for sections that need it
            if section_key == 'current_admission':
                target_layout.update()
                parent_widget.updateGeometry()
                parent_widget.update()
                print(f"[TRIBUNAL] Forced layout update for current_admission")

        except Exception as e:
            print(f"[TRIBUNAL] Failed to add CollapsibleSection to '{section_key}': {e}")
            import traceback
            traceback.print_exc()

    def _on_imported_checkbox_toggled(self, section_key: str, text: str, checked: bool):
        """Handle when an imported data checkbox is toggled."""
        if section_key not in self.cards:
            print(f"[TRIBUNAL] Warning: section_key '{section_key}' not in cards")
            return

        card = self.cards[section_key]
        if not hasattr(card, 'editor'):
            return

        current_text = card.editor.toPlainText()

        if checked:
            # Add text to card if not already there
            if text not in current_text:
                if current_text.strip():
                    new_text = current_text.strip() + "\n\n" + text
                else:
                    new_text = text
                card.editor.setPlainText(new_text)
                print(f"[TRIBUNAL] Added imported text to '{section_key}'")
        else:
            # Remove text from card
            if text in current_text:
                new_text = current_text.replace(text, "").strip()
                # Clean up double newlines
                while "\n\n\n" in new_text:
                    new_text = new_text.replace("\n\n\n", "\n\n")
                card.editor.setPlainText(new_text)
                print(f"[TRIBUNAL] Removed imported text from '{section_key}'")

    def _update_imported_checkbox_states(self, section_key: str):
        """Update imported checkbox states based on current card content."""
        if section_key not in self.popups:
            return

        popup = self.popups[section_key]
        if not hasattr(popup, '_imported_checkboxes'):
            return

        card_text = ""
        if section_key in self.cards:
            card_text = self.cards[section_key].editor.toPlainText().lower()

        for cb in popup._imported_checkboxes:
            full_text = cb.property("full_text") or ""
            line_lower = full_text.lower()
            # Check if significant words are in the card
            significant_words = [w for w in line_lower.split()[:6] if len(w) > 3]
            is_in_card = any(word in card_text for word in significant_words) if significant_words else False
            cb.blockSignals(True)
            cb.setChecked(is_in_card)
            cb.blockSignals(False)

    def _populate_diagnosis_popup(self, popup, content: str):
        """Populate diagnosis popup - try to match ICD-10 codes from imported report."""
        import re

        # Check for Yes/No answer and set radio button
        content_lower = content.lower().strip()
        is_yes = content_lower.startswith("yes") or "mental disorder" in content_lower or re.search(r'[Ff]\d{2}', content)
        is_no = content_lower.startswith("no") and not is_yes

        if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
            if is_yes:
                popup.yes_btn.setChecked(True)
                if hasattr(popup, 'diagnosis_container'):
                    popup.diagnosis_container.setVisible(True)
                print(f"[TRIBUNAL] Set diagnosis to Yes")
            elif is_no:
                popup.no_btn.setChecked(True)
                if hasattr(popup, 'diagnosis_container'):
                    popup.diagnosis_container.setVisible(False)
                print(f"[TRIBUNAL] Set diagnosis to No")
                return  # No need to search for diagnoses if No

        # Search ALL imported report sections for diagnoses, not just the diagnosis section
        all_content = content
        if hasattr(self, '_imported_report_data') and self._imported_report_data:
            # Combine all imported sections to search for diagnoses
            all_content = "\n".join(self._imported_report_data.values())
            print(f"[TRIBUNAL] Searching {len(self._imported_report_data)} imported sections for diagnoses")

        print(f"[TRIBUNAL] _populate_diagnosis_popup searching in {len(all_content)} chars")

        # Extract diagnosis text
        detail = all_content
        if content.lower().startswith("yes"):
            for prefix in ["Yes - ", "Yes, ", "Yes-", "yes - ", "yes, ", "yes-"]:
                if content.startswith(prefix):
                    detail = content[len(prefix):].strip() + "\n" + all_content
                    break

        if not hasattr(popup, 'dx_boxes') or not popup.dx_boxes:
            print("[TRIBUNAL] Diagnosis popup has no dx_boxes")
            return

        # ============================================================
        # Category-based diagnosis extraction (matching iOS logic)
        # Each pattern has a category — only the FIRST match per
        # category is kept, and max 3 diagnoses total.
        # Patterns are ordered specific-first so subtypes win over
        # generic terms within the same category.
        # ============================================================
        # (regex, search_term, category)
        diagnosis_patterns = [
            # Schizophrenia (specific subtypes first)
            (r'\bparanoid schizophrenia\b', "paranoid schizophrenia", "schizophrenia"),
            (r'\bcatatonic schizophrenia\b', "catatonic schizophrenia", "schizophrenia"),
            (r'\bhebephrenic schizophrenia\b', "hebephrenic schizophrenia", "schizophrenia"),
            (r'\bresidual schizophrenia\b', "residual schizophrenia", "schizophrenia"),
            (r'\bsimple schizophrenia\b', "simple schizophrenia", "schizophrenia"),
            (r'\bundifferentiated schizophrenia\b', "undifferentiated schizophrenia", "schizophrenia"),
            (r'\bschizophrenia\b', "schizophrenia", "schizophrenia"),
            # Schizoaffective (own category)
            (r'\bschizoaffective\b', "schizoaffective", "schizoaffective"),
            # Bipolar / mania
            (r'\bbipolar\s+affective\s+disorder\b', "bipolar", "bipolar"),
            (r'\bbipolar\s+disorder\b', "bipolar", "bipolar"),
            (r'\bbipolar\b', "bipolar", "bipolar"),
            (r'\bmanic\s+depressi', "bipolar", "bipolar"),
            # Depression
            (r'\brecurrent\s+depressi', "recurrent depressive", "depression"),
            (r'\bmajor\s+depressi', "depression", "depression"),
            (r'\bdepressive\s+episode\b', "depressive", "depression"),
            (r'\bdepressi', "depression", "depression"),
            # Personality disorders
            (r'\bemotionally\s+unstable\s+personality\b', "emotionally unstable personality", "personality"),
            (r'\bEUPD\b', "emotionally unstable personality", "personality"),
            (r'\bborderline\s+personality\b', "borderline", "personality"),
            (r'\bantisocial\s+personality\b', "antisocial personality", "personality"),
            (r'\bdissocial\s+personality\b', "antisocial personality", "personality"),
            (r'\bASPD\b', "antisocial personality", "personality"),
            (r'\bnarcissistic\s+personality\b', "narcissistic personality", "personality"),
            (r'\bparanoid\s+personality\b', "paranoid personality", "personality"),
            (r'\bpersonality\s+disorder\b', "personality disorder", "personality"),
            # Anxiety
            (r'\bPTSD\b', "PTSD", "anxiety"),
            (r'\bpost[- ]traumatic\s+stress\b', "PTSD", "anxiety"),
            (r'\bOCD\b', "obsessive", "anxiety"),
            (r'\bobsessive[- ]compulsive\s+disorder\b', "obsessive", "anxiety"),
            (r'\bgeneralised\s+anxiety\b', "anxiety", "anxiety"),
            (r'\bgeneralized\s+anxiety\b', "anxiety", "anxiety"),
            # Psychosis (only if no schizophrenia/schizoaffective)
            (r'\bpsychosis\b', "psychosis", "psychosis"),
            (r'\bpsychotic\s+disorder\b', "psychotic", "psychosis"),
            # Autism
            (r'\bautism\s+spectrum\s+disorder\b', "autism", "autism"),
            (r'\bautistic\s+spectrum\b', "autism", "autism"),
            (r'\basperger\b', "autism", "autism"),
            (r'\batypical\s+autism\b', "autism", "autism"),
            (r'\bASD\b', "autism", "autism"),
            # Learning disability
            (r'\blearning\s+disabilit', "learning disability", "learning_disability"),
            (r'\bintellectual\s+disabilit', "learning disability", "learning_disability"),
            # ADHD
            (r'\bADHD\b', "ADHD", "adhd"),
            (r'\battention\s+deficit\b', "ADHD", "adhd"),
            # Substance
            (r'\balcohol\s+dependence\b', "alcohol", "substance"),
            (r'\bdrug\s+dependence\b', "drug", "substance"),
            (r'\bopioid\s+dependence\b', "opioid", "substance"),
            (r'\bsubstance\s+misuse\b', "substance", "substance"),
            (r'\bsubstance\s+use\s+disorder\b', "substance", "substance"),
        ]

        # Diagnostic hierarchy: finding a higher-level diagnosis suppresses lower ones
        # e.g. schizophrenia suppresses psychosis; schizoaffective suppresses psychosis
        HIERARCHY_SUPPRESSES = {
            "schizophrenia": {"psychosis"},
            "schizoaffective": {"psychosis"},
            "bipolar": {"depression"},
        }

        diagnoses_to_set = []
        seen_categories = set()

        for pattern, keyword, category in diagnosis_patterns:
            if category in seen_categories:
                continue
            if re.search(pattern, detail, re.IGNORECASE):
                seen_categories.add(category)
                # Also suppress lower-hierarchy categories
                suppressed = HIERARCHY_SUPPRESSES.get(category, set())
                seen_categories.update(suppressed)
                diagnoses_to_set.append({"type": "name", "term": keyword, "category": category})
                print(f"[TRIBUNAL] Diagnosis match: '{keyword}' (category={category})")
                if suppressed:
                    print(f"[TRIBUNAL] Suppressed categories: {suppressed}")
                if len(diagnoses_to_set) >= 3:
                    break

        # Also check for ICD-10 codes not already covered by a category match
        icd_category_map = {
            "F20": "schizophrenia", "F21": "schizophrenia",
            "F22": "psychosis", "F23": "psychosis",
            "F25": "schizoaffective",
            "F30": "bipolar", "F31": "bipolar",
            "F32": "depression", "F33": "depression",
            "F40": "anxiety", "F41": "anxiety", "F42": "anxiety", "F43": "anxiety",
            "F60": "personality",
            "F70": "learning_disability", "F71": "learning_disability",
            "F72": "learning_disability", "F73": "learning_disability",
            "F84": "autism",
            "F10": "substance", "F11": "substance", "F12": "substance",
            "F13": "substance", "F14": "substance", "F19": "substance",
        }
        if len(diagnoses_to_set) < 3:
            icd_pattern = r'([Ff]\d{2}(?:\.\d{1,2})?)'
            for match in re.finditer(icd_pattern, detail):
                code = match.group(1).upper()
                code_prefix = code[:3]
                cat = icd_category_map.get(code_prefix)
                if cat and cat not in seen_categories:
                    seen_categories.add(cat)
                    diagnoses_to_set.append({"type": "icd", "term": code, "category": cat})
                    print(f"[TRIBUNAL] ICD-10 match: {code} (category={cat})")
                    if len(diagnoses_to_set) >= 3:
                        break

        print(f"[TRIBUNAL] Final diagnoses to match: {diagnoses_to_set}")

        # Set combo boxes for found diagnoses (up to 3)
        combo_idx = 0
        for dx_info in diagnoses_to_set:
            if combo_idx >= len(popup.dx_boxes):
                break

            combo = popup.dx_boxes[combo_idx]
            search_term = dx_info["term"].lower()
            found = False
            fallback_idx = -1

            # Search combo items for match
            for j in range(combo.count()):
                item_text = combo.itemText(j).lower()

                # Check if search term is in item text
                if search_term in item_text:
                    # Check for negative phrases before the search term
                    term_pos = item_text.find(search_term)
                    prefix = item_text[:term_pos] if term_pos > 0 else ""
                    if "no symptoms of" in prefix or "without" in prefix or "absence of" in prefix:
                        print(f"[TRIBUNAL] Skipping negative match: {combo.itemText(j)}")
                        continue

                    # Prefer items that START with the search term
                    if item_text.startswith(search_term):
                        combo.setCurrentIndex(j)
                        print(f"[TRIBUNAL] Set diagnosis {combo_idx+1} to: {combo.itemText(j)} (exact)")
                        found = True
                        combo_idx += 1
                        break
                    elif fallback_idx < 0:
                        fallback_idx = j  # First valid contains match as fallback

                # Also check if ICD code is in item text (in parentheses)
                if dx_info["type"] == "icd":
                    if f"({search_term})" in item_text or f"({search_term.lower()})" in item_text:
                        combo.setCurrentIndex(j)
                        print(f"[TRIBUNAL] Set diagnosis {combo_idx+1} to: {combo.itemText(j)} (ICD)")
                        found = True
                        combo_idx += 1
                        break

            # Fallback to contains match if no startswith match found
            if not found and fallback_idx >= 0:
                combo.setCurrentIndex(fallback_idx)
                print(f"[TRIBUNAL] Set diagnosis {combo_idx+1} to: {combo.itemText(fallback_idx)} (contains)")
                found = True
                combo_idx += 1

            if not found:
                print(f"[TRIBUNAL] Could not find match for: {dx_info['term']}")

        # If we found any diagnoses, auto-click Yes and show diagnosis container
        if combo_idx > 0:
            if hasattr(popup, 'yes_btn') and not popup.yes_btn.isChecked():
                popup.yes_btn.setChecked(True)
                print(f"[TRIBUNAL] Auto-clicked Yes because diagnoses were found")
            if hasattr(popup, 'diagnosis_container'):
                popup.diagnosis_container.setVisible(True)

        print(f"[TRIBUNAL] Populated diagnosis popup with {combo_idx} entries")

    def _populate_treatment_popup(self, popup, content: str):
        """Parse medication list from imported content and populate popup medication entries.

        Uses CANONICAL_MEDS dictionary for comprehensive medication matching (same as nursing section 4).
        """
        import re
        from CANONICAL_MEDS import MEDICATIONS

        if not content:
            return

        print(f"[TRIBUNAL] Parsing medications from imported content...")

        # Frequency mapping
        FREQ_MAP = {
            "od": "OD", "once daily": "OD", "daily": "OD", "mane": "OD", "nocte": "Nocte",
            "bd": "BD", "twice daily": "BD", "twice a day": "BD",
            "tds": "TDS", "three times daily": "TDS", "three times a day": "TDS",
            "qds": "QDS", "four times daily": "QDS", "four times a day": "QDS",
            "prn": "PRN", "as required": "PRN", "when required": "PRN",
            "weekly": "Weekly", "once weekly": "Weekly", "1 weekly": "Weekly",
            "fortnightly": "Fortnightly", "2 weekly": "Fortnightly", "every 2 weeks": "Fortnightly",
            "3 weekly": "3 Weekly", "every 3 weeks": "3 Weekly",
            "monthly": "Monthly", "4 weekly": "Monthly", "every 4 weeks": "Monthly",
        }

        # Parse medication lines - look for patterns like:
        # - Olanzapine 10mg OD
        # • Depakote 500mg BD
        # Sertraline 100mg once daily
        lines = content.split('\n')
        meds_found = []

        # Build a lowercase -> uppercase key map for matching
        med_name_map = {}
        for med_key, med_info in MEDICATIONS.items():
            med_name_map[med_key.lower()] = med_key
            canonical = med_info.get("canonical", "").lower()
            if canonical:
                med_name_map[canonical] = med_key
            for pattern in med_info.get("patterns", []):
                med_name_map[pattern.lower()] = med_key

        for line in lines:
            line = line.strip()
            # Remove bullet points and list markers
            line = re.sub(r'^[\-\•\*\d\.]+\s*', '', line)
            if not line:
                continue

            # Try to match medication name
            words = line.split()
            med_key = None
            med_name = None

            # Check first 1-3 words for medication name
            for i in range(1, min(4, len(words) + 1)):
                test_name = ' '.join(words[:i]).lower()
                # Remove trailing punctuation
                test_name = re.sub(r'[,.:;]+$', '', test_name)
                if test_name in med_name_map:
                    med_key = med_name_map[test_name]
                    med_name = ' '.join(words[:i])
                    break

            if not med_key:
                # Try matching against patterns in the full line
                line_lower = line.lower()
                for name, key in med_name_map.items():
                    if re.search(r'\b' + re.escape(name) + r'\b', line_lower):
                        med_key = key
                        med_name = name
                        break

            if not med_key:
                continue

            # Extract dose (number followed by mg, mcg, g, ml, etc.)
            dose_match = re.search(r'(\d+(?:\.\d+)?)\s*(mg|mcg|g|ml|micrograms?)', line, re.IGNORECASE)
            dose = dose_match.group(0) if dose_match else ""

            # Extract frequency
            freq = "OD"  # Default
            line_lower = line.lower()
            for freq_pattern, freq_value in FREQ_MAP.items():
                if freq_pattern in line_lower:
                    freq = freq_value
                    break

            meds_found.append({
                "med_key": med_key,
                "dose": dose,
                "freq": freq
            })
            print(f"[TRIBUNAL] Found medication: {med_key} {dose} {freq}")

        if not meds_found:
            print(f"[TRIBUNAL] No medications found in imported content")
            return

        # Clear existing empty medication entries (keep only the first one if empty)
        while len(popup._medications) > 1:
            entry = popup._medications[-1]
            if entry["name"].currentText().strip() == "":
                popup._medications.remove(entry)
                entry["widget"].deleteLater()
            else:
                break

        # Add enough entries
        while len(popup._medications) < len(meds_found):
            popup._add_medication_entry()

        # Populate the medication entries
        for i, med in enumerate(meds_found):
            if i >= len(popup._medications):
                break

            entry = popup._medications[i]
            med_key = med.get("med_key", "")

            # Set medication name
            name_combo = entry.get("name")
            if name_combo and med_key:
                idx = name_combo.findText(med_key)
                if idx >= 0:
                    name_combo.setCurrentIndex(idx)
                else:
                    # Set as text (editable combo)
                    name_combo.setCurrentText(med_key)

            # Set dose
            dose_combo = entry.get("dose")
            if dose_combo and med.get("dose"):
                dose_str = med.get("dose")
                idx = dose_combo.findText(dose_str)
                if idx >= 0:
                    dose_combo.setCurrentIndex(idx)
                else:
                    dose_combo.setCurrentText(dose_str)

            # Set frequency
            freq_combo = entry.get("freq")
            if freq_combo and med.get("freq"):
                freq = med.get("freq")
                idx = freq_combo.findText(freq)
                if idx >= 0:
                    freq_combo.setCurrentIndex(idx)

        print(f"[TRIBUNAL] ✓ Pre-filled {len(meds_found)} medication(s) from import")

        # Set imported data with checkboxes
        if hasattr(popup, 'set_imported_data'):
            # Get current card content to check what's already there
            card_text = ""
            if hasattr(self, 'cards') and 'treatment' in self.cards:
                card_text = self.cards['treatment'].editor.toPlainText()
            popup.set_imported_data(content, card_text)
            print(f"[TRIBUNAL] Set imported data for treatment popup")

    def _populate_current_admission_extracted_section(self, popup, content: str):
        """Populate current_admission popup's built-in extracted_section with imported text."""
        from PySide6.QtWidgets import QTextEdit
        from PySide6.QtCore import Qt

        print(f"[TRIBUNAL] _populate_current_admission_extracted_section called")
        print(f"[TRIBUNAL]   has extracted_section: {hasattr(popup, 'extracted_section')}")
        print(f"[TRIBUNAL]   has extracted_checkboxes_layout: {hasattr(popup, 'extracted_checkboxes_layout')}")
        print(f"[TRIBUNAL]   content to add: '{content[:100]}...' ({len(content)} chars)")

        if not hasattr(popup, 'extracted_section') or not popup.extracted_section:
            print(f"[TRIBUNAL] Current admission popup has no extracted_section")
            return

        if not content or not content.strip():
            print(f"[TRIBUNAL] Current admission: no content to populate")
            return

        try:
            # Clear existing content in the extracted_checkboxes_layout
            if hasattr(popup, 'extracted_checkboxes_layout'):
                layout = popup.extracted_checkboxes_layout
                print(f"[TRIBUNAL] Clearing layout with {layout.count()} items")
                # Clear existing widgets
                while layout.count():
                    item = layout.takeAt(0)
                    if item.widget():
                        item.widget().deleteLater()

                # Add content as a read-only QTextEdit (shows full content with scroll)
                content_edit = QTextEdit()
                content_edit.setPlainText(content)
                content_edit.setReadOnly(True)
                content_edit.setMinimumHeight(120)  # Reduced by 20%
                content_edit.setStyleSheet("""
                    QTextEdit {
                        font-size: 16px;
                        color: #4a4a4a;
                        background: rgba(255, 248, 220, 0.95);
                        border: 1px solid rgba(180, 150, 50, 0.3);
                        border-radius: 6px;
                        padding: 8px;
                    }
                """)
                layout.addWidget(content_edit)
                print(f"[TRIBUNAL] Added QTextEdit widget to layout")

            # Show the section and expand it to show content
            popup.extracted_section.setVisible(True)
            print(f"[TRIBUNAL] Set extracted_section visible")
            # Set a larger content height if method available
            if hasattr(popup.extracted_section, 'set_content_height'):
                popup.extracted_section.set_content_height(300)

            # Force expand the section - directly set state and show components
            section = popup.extracted_section
            if hasattr(section, '_is_collapsed'):
                was_collapsed = section._is_collapsed
                section._is_collapsed = False
                if hasattr(section, 'collapse_btn'):
                    section.collapse_btn.setText("−")
                if hasattr(section, 'content_container'):
                    section.content_container.setVisible(True)
                if hasattr(section, 'drag_bar'):
                    section.drag_bar.setVisible(True)
                if hasattr(section, '_update_height'):
                    section._update_height()
                print(f"[TRIBUNAL] Force-expanded current_admission extracted_section (was_collapsed={was_collapsed})")

            popup._imported_data_added = True
            print(f"[TRIBUNAL] Populated current_admission extracted_section with imported text ({len(content)} chars)")

        except Exception as e:
            print(f"[TRIBUNAL] Failed to populate current_admission extracted_section: {e}")
            import traceback
            traceback.print_exc()

    def _populate_forensic_extracted_section(self, popup, content: str):
        """Populate forensic popup's built-in extracted_section with imported text."""
        from PySide6.QtWidgets import QTextEdit
        from PySide6.QtCore import Qt

        if not hasattr(popup, 'extracted_section') or not popup.extracted_section:
            print(f"[TRIBUNAL] Forensic popup has no extracted_section")
            return

        if not content or not content.strip():
            return

        try:
            # Clear existing content in the extracted_checkboxes_layout
            if hasattr(popup, 'extracted_checkboxes_layout'):
                layout = popup.extracted_checkboxes_layout
                # Clear existing widgets
                while layout.count():
                    item = layout.takeAt(0)
                    if item.widget():
                        item.widget().deleteLater()

                # Add content as a read-only QTextEdit (shows full content with scroll)
                content_edit = QTextEdit()
                content_edit.setPlainText(content)
                content_edit.setReadOnly(True)
                content_edit.setMinimumHeight(120)  # Reduced by 20%
                content_edit.setStyleSheet("""
                    QTextEdit {
                        font-size: 15px;
                        color: #4a4a4a;
                        background: transparent;
                        border: none;
                    }
                """)
                layout.addWidget(content_edit)

            # Store imported text for formatted_text() to use
            popup._imported_report_text = content

            # Show the section and expand it to show content
            popup.extracted_section.setVisible(True)
            # Set a larger content height if method available
            if hasattr(popup.extracted_section, 'set_content_height'):
                popup.extracted_section.set_content_height(300)
            popup._imported_data_added = True
            print(f"[TRIBUNAL] Populated forensic extracted_section with full imported text ({len(content)} chars)")

        except Exception as e:
            print(f"[TRIBUNAL] Failed to populate forensic extracted_section: {e}")
            import traceback
            traceback.print_exc()

    def _populate_psych_history_imported(self, popup, sections_list: list):
        """Populate previous_mh_dates popup with separate text widgets per report category."""
        from PySide6.QtWidgets import QTextEdit, QLabel, QFrame
        from PySide6.QtCore import QEvent

        layout = popup.extracted_checkboxes_layout
        # Clear existing
        while layout.count():
            item = layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        # Store combined text for the include checkbox
        all_texts = []
        text_edits = []
        for cat_name, text in sections_list:
            all_texts.append(text)

            # Section label
            label = QLabel(cat_name)
            label.setStyleSheet("""
                QLabel {
                    font-size: 17px;
                    font-weight: 600;
                    color: #806000;
                    background: transparent;
                    border: none;
                    padding-top: 4px;
                }
            """)
            layout.addWidget(label)

            # Text widget
            text_edit = QTextEdit()
            text_edit.setPlainText(text)
            text_edit.setReadOnly(True)
            text_edit.setMinimumHeight(100)
            text_edit.setMaximumHeight(250)
            text_edit.setStyleSheet("""
                QTextEdit {
                    font-size: 16px;
                    color: #4a4a4a;
                    background: rgba(255, 255, 255, 0.8);
                    border: 1px solid rgba(180, 150, 50, 0.3);
                    border-radius: 6px;
                    padding: 6px;
                }
            """)
            layout.addWidget(text_edit)
            text_edits.append(text_edit)

        # Add drag handle INSIDE the content area (resizes text edits only)
        drag_handle = QFrame()
        drag_handle.setFixedHeight(10)
        drag_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        drag_handle.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(180,150,50,0.1), stop:0.5 rgba(180,150,50,0.35), stop:1 rgba(180,150,50,0.1));
                border-radius: 2px;
                margin: 2px 60px;
            }
            QFrame:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 rgba(180,150,50,0.2), stop:0.5 rgba(180,150,50,0.6), stop:1 rgba(180,150,50,0.2));
            }
        """)
        layout.addWidget(drag_handle)

        # Drag handle state
        drag_handle._drag_y = None
        drag_handle._init_heights = None

        def _on_press(ev, h=drag_handle, edits=text_edits):
            h._drag_y = ev.globalPosition().y()
            h._init_heights = [e.height() for e in edits]

        def _on_move(ev, h=drag_handle, edits=text_edits):
            if h._drag_y is not None:
                delta = int(ev.globalPosition().y() - h._drag_y)
                for i, e in enumerate(edits):
                    new_h = max(80, h._init_heights[i] + delta)
                    e.setMinimumHeight(new_h)
                    e.setMaximumHeight(new_h)

        def _on_release(ev, h=drag_handle, edits=text_edits):
            if h._drag_y is not None:
                for e in edits:
                    e.setMaximumHeight(16777215)
                h._drag_y = None

        drag_handle.mousePressEvent = _on_press
        drag_handle.mouseMoveEvent = _on_move
        drag_handle.mouseReleaseEvent = _on_release

        # Hide the built-in drag bar and remove fixed height so section sizes naturally
        sec = popup.extracted_section
        sec.drag_bar.setVisible(False)
        sec.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        sec.setMinimumHeight(0)
        sec.setMaximumHeight(16777215)

        popup._imported_report_text = "\n\n".join(all_texts)
        sec.setVisible(True)
        # Don't use set_content_height (it calls setFixedHeight) — let content size naturally
        sec._is_collapsed = False
        sec.collapse_btn.setText("−")
        sec.content_container.setVisible(True)
        popup._imported_data_added = True
        print(f"[TRIBUNAL] Populated previous_mh_dates with {len(sections_list)} separate imported sections")

    def _populate_mca_dol_popup(self, popup, content: str):
        """Populate MCA/DoL popup based on imported content."""
        content_lower = content.lower()

        # DoLsPopup has yes_btn, no_btn, na_btn radio buttons
        if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
            # Check for DoLs keywords
            if "dols is required" in content_lower or "dol is required" in content_lower:
                popup.yes_btn.setChecked(True)
                print(f"[TRIBUNAL] Set mca_dol popup to Yes (DoLs required)")
            elif "dols is not required" in content_lower or "dol is not required" in content_lower:
                popup.no_btn.setChecked(True)
                print(f"[TRIBUNAL] Set mca_dol popup to No (DoLs not required)")
            elif "not applicable" in content_lower or "n/a" in content_lower:
                if hasattr(popup, 'na_btn'):
                    popup.na_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set mca_dol popup to N/A")
            # Also check for implicit indicators
            elif "not likely" in content_lower or "would not be" in content_lower:
                popup.no_btn.setChecked(True)
                print(f"[TRIBUNAL] Set mca_dol popup to No (inferred)")
            elif "lacks capacity" in content_lower or "deprivation of liberty" in content_lower:
                popup.yes_btn.setChecked(True)
                print(f"[TRIBUNAL] Set mca_dol popup to Yes (inferred)")

        print(f"[TRIBUNAL] Populated MCA/DoL popup")

    def _populate_other_detention_popup(self, popup, content: str):
        """Populate other detention popup based on imported content with x marks."""
        import re

        def has_yes_cross(text):
            """Check if Yes has a checked box symbol."""
            if re.search(r'yes\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            if re.search(r'yes\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes', text, re.IGNORECASE):
                return True
            # Pattern: Yes checked, No unchecked
            if re.search(r'yes\s*[☒☑✓✔\[xX\]].*no\s*[☐□\[\s\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*yes.*[☐□]\s*no', text, re.IGNORECASE):
                return True
            return False

        def has_no_cross(text):
            """Check if No has a checked box symbol (meaning answer is No)."""
            if re.search(r'no\s*\[\s*[xX]\s*\]', text, re.IGNORECASE):
                return True
            if re.search(r'no\s*[☒☑✓✔]', text, re.IGNORECASE):
                return True
            if re.search(r'[☒☑✓✔]\s*no\b', text, re.IGNORECASE):
                return True
            # Pattern: Yes unchecked, No checked
            if re.search(r'yes\s*[☐□\[\s\]].*no\s*[☒☑✓✔\[xX\]]', text, re.IGNORECASE):
                return True
            if re.search(r'[☐□]\s*yes.*[☒☑✓✔]\s*no', text, re.IGNORECASE):
                return True
            return False

        content_lower = content.lower()
        print(f"[TRIBUNAL] Section 20 content: {repr(content[:200] if len(content) > 200 else content)}")

        if hasattr(popup, 'yes_btn') and hasattr(popup, 'no_btn'):
            yes_found = has_yes_cross(content)
            no_found = has_no_cross(content)
            print(f"[TRIBUNAL] Section 20 detection: yes_cross={yes_found}, no_cross={no_found}")
            # Check for explicit crosses first
            if yes_found:
                popup.yes_btn.setChecked(True)
                print(f"[TRIBUNAL] Set other_detention to Yes (found Yes [x])")
            elif no_found:
                popup.no_btn.setChecked(True)
                print(f"[TRIBUNAL] Set other_detention to No (found No [x])")
            elif "n/a" in content_lower or "not applicable" in content_lower:
                if hasattr(popup, 'na_btn'):
                    popup.na_btn.setChecked(True)
                    print(f"[TRIBUNAL] Set other_detention to N/A")
            # Fallback to text-based detection
            elif content_lower.startswith("yes"):
                popup.yes_btn.setChecked(True)
                print(f"[TRIBUNAL] Set other_detention to Yes (text-based)")
            elif content_lower.startswith("no"):
                popup.no_btn.setChecked(True)
                print(f"[TRIBUNAL] Set other_detention to No (text-based)")

    def _clear_report(self):
        """Clear all cards to start a new report."""
        from PySide6.QtWidgets import QMessageBox

        # Confirm with user
        reply = QMessageBox.question(
            self,
            "Clear Report",
            "Are you sure you want to clear all content and start a new report?\n\nThis cannot be undone.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            # Clear all card editors
            for key, card in self.cards.items():
                if hasattr(card, 'editor'):
                    card.editor.clear()

            # Clear stored extracted data
            self._extracted_raw_notes = []
            self._extracted_categories = {}
            self._incident_data = []
            if hasattr(self, '_imported_report_data'):
                self._imported_report_data = {}
            if hasattr(self, '_imported_report_sections'):
                self._imported_report_sections = {}

            # Destroy all popups and remove from stack
            for key, popup in list(self.popups.items()):
                self.popup_stack.removeWidget(popup)
                popup.deleteLater()
            self.popups.clear()

            # Clear data extractor
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                if hasattr(self._data_extractor_overlay, 'clear_extraction'):
                    self._data_extractor_overlay.clear_extraction()

            # Recreate author/signature popups so mydetails get restored
            for restore_key in ("author", "signature"):
                self._on_card_clicked(restore_key)

            print("[TRIBUNAL] Report cleared - ready for new report")

    def _open_data_extractor_overlay(self):
        """Create the data extractor (hidden) for background processing."""
        from data_extractor_popup import DataExtractorPopup

        # Create data extractor if not exists
        if not hasattr(self, '_data_extractor_overlay') or not self._data_extractor_overlay:
            self._data_extractor_overlay = DataExtractorPopup(parent=self)
            self._data_extractor_overlay.hide()

            # Connect the data extraction signal to populate fixed panels
            if hasattr(self._data_extractor_overlay, 'data_extracted'):
                self._data_extractor_overlay.data_extracted.connect(self._on_data_extracted)

    def _close_data_extractor_overlay(self):
        """Close the data extractor overlay and restore the previous popup."""
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            self._data_extractor_overlay.hide()

            # Restore the hidden popup if any
            if hasattr(self, '_hidden_popup') and self._hidden_popup:
                self._hidden_popup.show()
                self._hidden_popup = None

    def _on_data_extracted(self, data: dict):
        """Handle extracted data from the data extractor and populate fixed panels."""
        import os
        print(f"[TRIBUNAL] Data extracted: {list(data.keys())}")
        cov = data.get("_coverage")
        if cov and cov.get("uncategorised", 0) > 0:
            print(f"[TRIBUNAL] Warning: {cov['uncategorised']} paragraphs uncategorised "
                  f"({cov['categorised']}/{cov['total_paragraphs']} categorised)")

        # Check if this is filtered data to send to the current card
        filtered_category = data.get("filtered_category")
        if filtered_category and hasattr(self, '_selected_card_key') and self._selected_card_key:
            print(f"[TRIBUNAL] Filtered category '{filtered_category}' -> sending to current card '{self._selected_card_key}'")
            self._send_filtered_to_current_card(data)
            return

        # Skip if this exact data was already processed
        categories = data.get("categories", {})
        cat_keys = tuple(sorted(categories.keys())) if categories else ()
        cat_count = sum(len(v.get("items", [])) if isinstance(v, dict) else 0 for v in categories.values())
        content_sig = (cat_keys, cat_count)
        if self._data_processed_id == content_sig:
            print(f"[TRIBUNAL] Skipping _on_data_extracted - data already processed")
            return
        self._data_processed_id = content_sig

        # Check if data came from a report (not notes)
        is_report = False
        if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
            dtype = getattr(self._data_extractor_overlay, '_detected_document_type', None)
            if dtype == "reports":
                is_report = True
                print(f"[TRIBUNAL] Detected report data (dtype={dtype})")

        categories = data.get("categories", {})
        print(f"[TRIBUNAL] Available categories: {list(categories.keys())}")

        if is_report and categories:
            # Report pipeline: map categories directly to card sections
            source = os.path.basename(getattr(self, '_data_extractor_source_file', '') or '') or "Data Extractor"
            self._populate_from_report_categories(categories, source_filename=source)
        else:
            # Notes pipeline - skip if report data already imported (prevents cross-talk)
            if self._has_report_data():
                print(f"[TRIBUNAL] Skipping notes pipeline - report data already imported")
                return

            raw_notes = []
            if hasattr(self, '_data_extractor_overlay') and self._data_extractor_overlay:
                raw_notes = getattr(self._data_extractor_overlay, 'notes', [])

            # Fall back to SharedDataStore if no local notes
            if not raw_notes:
                try:
                    from shared_data_store import get_shared_store
                    shared_store = get_shared_store()
                    if shared_store.has_notes():
                        raw_notes = shared_store.notes
                        print(f"[TRIBUNAL] Got {len(raw_notes)} notes from SharedDataStore (global import)")
                except Exception as e:
                    print(f"[TRIBUNAL] Error getting notes from SharedDataStore: {e}")

            print(f"[TRIBUNAL] Raw notes available: {len(raw_notes)}")

            # Ask add/replace if existing notes
            action = self._ask_import_action("", "notes")
            if action == "cancel":
                return
            if action == "add":
                existing_notes = getattr(self, '_extracted_raw_notes', []) or []
                raw_notes = existing_notes + raw_notes
                existing_cats = getattr(self, '_extracted_categories', {}) or {}
                for cat_name, cat_data in categories.items():
                    if cat_name in existing_cats and isinstance(existing_cats[cat_name], dict) and isinstance(cat_data, dict):
                        existing_items = existing_cats[cat_name].get("items", [])
                        new_items = cat_data.get("items", [])
                        existing_cats[cat_name]["items"] = existing_items + new_items
                    else:
                        existing_cats[cat_name] = cat_data
                categories = existing_cats

            # STORE at page level for popups created later
            self._extracted_raw_notes = raw_notes
            self._extracted_categories = categories
            print(f"[TRIBUNAL] Stored {len(raw_notes)} raw notes and {len(categories)} categories at page level")

            # Populate any existing popups
            self._populate_fixed_panels()

    # Mapping from data extractor category names to tribunal card keys
    CATEGORY_TO_CARD = {
        "Forensic History": "forensic",
        "FORENSIC": "forensic",
        "Past Psychiatric History": "previous_mh_dates",
        "PAST_PSYCH": "previous_mh_dates",
        "Psychiatric History": "previous_mh_dates",
        "Circumstances of Admission": "current_admission",
        "HISTORY_OF_PRESENTING_COMPLAINT": "current_admission",
        "History of Presenting Complaint": "progress",
        "Diagnosis": "diagnosis",
        "Medication History": "treatment",
        "Strengths": "strengths",
        "Risk": "risk_harm",
        "SUMMARY": "recommendations",
        "Summary": "recommendations",
        "MENTAL_STATE": "progress",
        "Mental State": "progress",
        "BACKGROUND_HISTORY": "current_admission",
        "Background History": "current_admission",
        "SOCIAL_HISTORY": "community",
        "Social History": "community",
    }

    def _populate_from_report_categories(self, categories: dict, source_filename: str = ""):
        """Populate report sections from data extractor categories (report mode).

        Maps extracted category names to tribunal card keys and populates
        card editors directly, similar to _populate_from_docx().
        """
        from PySide6.QtWidgets import QMessageBox
        from shared_data_store import get_shared_store

        if not hasattr(self, '_imported_report_data'):
            self._imported_report_data = {}

        source_label = source_filename or "Data Extractor"
        action = self._ask_import_action(source_label, "report")
        if action == "cancel":
            return
        if action == "replace":
            self._imported_report_data.clear()

        # Build card_key -> combined text from categories
        card_texts = {}
        # Also track individual category contributions per card for popups
        if not hasattr(self, '_imported_report_sections'):
            self._imported_report_sections = {}

        # Build set of valid card keys for direct matching
        valid_card_keys = {key for _, key in self.SECTIONS}

        for cat_name, cat_data in categories.items():
            # Categories are now card keys directly (e.g. "forensic", "treatment")
            if cat_name in valid_card_keys:
                card_key = cat_name
            else:
                # Fallback to old category name mapping
                card_key = self.CATEGORY_TO_CARD.get(cat_name)
            if not card_key:
                print(f"[TRIBUNAL] No card mapping for category: {cat_name}")
                continue

            # Combine all items' text for this category
            items = cat_data.get("items", []) if isinstance(cat_data, dict) else []
            texts = []
            for item in items:
                if isinstance(item, dict):
                    text = item.get("text", "") or item.get("content", "")
                elif isinstance(item, str):
                    text = item
                else:
                    text = str(item)
                if text:
                    texts.append(text.strip())

            if not texts:
                continue

            combined = "\n\n".join(texts)
            # Append if multiple categories map to the same card
            if card_key in card_texts:
                card_texts[card_key] += "\n\n" + combined
            else:
                card_texts[card_key] = combined

            # Track separate sections per card
            if card_key not in self._imported_report_sections:
                self._imported_report_sections[card_key] = []
            self._imported_report_sections[card_key].append((cat_name, combined))

        # Store imported data (cards are NOT auto-filled; user sends from popups)
        sections_for_store = {}
        for card_key, content in card_texts.items():
            if action == "add":
                content = self._merge_report_section(card_key, content, source_label)
            self._imported_report_data[card_key] = content
            sections_for_store[card_key] = content
            print(f"[TRIBUNAL] Stored report category for section '{card_key}' (ready in popup)")

        # Populate popups with imported data so user can review and send
        self._populate_popups_from_import(sections_for_store)

        # Push patient details to shared store
        self._push_patient_details_to_shared_store(sections_for_store)

        # Push sections to shared store for cross-talk
        shared_store = get_shared_store()
        shared_store.set_report_sections(sections_for_store, source_form="tribunal")

        # Show success message
        mapped_count = len(sections_for_store)
        action_word = "Added" if action == "add" else "Loaded"
        QMessageBox.information(
            self,
            "Report Loaded",
            f"{action_word} report data from {source_label}.\n\n"
            f"{action_word} {mapped_count} sections to popups.\n\n"
            f"Click each card to review and send the content."
        )

        print(f"[TRIBUNAL] {action_word} {mapped_count} sections from report categories to popups")

    def _send_filtered_to_current_card(self, data: dict):
        """Send filtered data from data extractor to the currently selected card."""
        current_key = self._selected_card_key
        if not current_key:
            print("[TRIBUNAL] No card selected - cannot send filtered data")
            return

        categories = data.get("categories", {})
        filtered_category = data.get("filtered_category", "")

        # Combine all items from the filtered categories
        all_items = []
        for cat_name, cat_data in categories.items():
            if isinstance(cat_data, dict) and "items" in cat_data:
                all_items.extend(cat_data.get("items", []))

        if not all_items:
            print(f"[TRIBUNAL] No items found in filtered data for '{filtered_category}'")
            return

        # Format the text content
        texts = []
        for item in all_items:
            text = item.get("text", "").strip()
            date_str = item.get("date", "")
            if text:
                if date_str:
                    texts.append(f"[{date_str}] {text}")
                else:
                    texts.append(text)

        combined_text = "\n\n".join(texts)
        print(f"[TRIBUNAL] Sending {len(texts)} items to card '{current_key}'")

        # Update the card with the combined text
        if current_key in self.cards:
            self._update_card(current_key, combined_text)
            print(f"[TRIBUNAL] Card '{current_key}' updated with filtered data")

    def _populate_fixed_panels(self):
        """Populate all fixed panels with extracted data."""
        raw_notes = self._extracted_raw_notes
        categories = self._extracted_categories

        if not raw_notes:
            print("[TRIBUNAL] No extracted data to populate panels")
            return

        # Helper to get items from category (categories have "items" key)
        def get_category_items(cat_name):
            cat = categories.get(cat_name, {})
            if isinstance(cat, dict):
                return cat.get("items", [])
            return []

        # Section 7: Previous admissions (skip if report data imported)
        if "previous_admission_reasons" in self.popups and not (hasattr(self, '_imported_report_data') and 'previous_admission_reasons' in self._imported_report_data):
            popup = self.popups["previous_admission_reasons"]
            if hasattr(popup, 'notes'):
                popup.notes = raw_notes
            if hasattr(popup, 'set_entries'):
                psych_items = get_category_items("Past Psychiatric History")
                entries = psych_items if psych_items else raw_notes[:20]
                popup.set_entries(entries, f"{len(entries)} notes")
                print(f"[TRIBUNAL] Populated section 7 with {len(entries)} notes")
            # Run timeline analysis to detect admissions
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[TRIBUNAL] Ran timeline analysis for section 7")

        # Section 6: Past psychiatric history (skip if report data imported)
        if "previous_mh_dates" in self.popups and not (hasattr(self, '_imported_report_data') and 'previous_mh_dates' in self._imported_report_data):
            popup = self.popups["previous_mh_dates"]
            if hasattr(popup, 'notes'):
                popup.notes = raw_notes
            if hasattr(popup, 'set_entries'):
                psych_items = get_category_items("Past Psychiatric History")
                entries = psych_items if psych_items else raw_notes[:20]
                popup.set_entries(entries, f"{len(entries)} notes")
                print(f"[TRIBUNAL] Populated section 6 with {len(entries)} notes")
            # Run timeline analysis to detect admissions
            if hasattr(popup, 'set_notes'):
                popup.set_notes(raw_notes)
                print(f"[TRIBUNAL] Ran timeline analysis for section 6")

        # Section 8: Current admission (uses filtered notes)
        if "current_admission" in self.popups:
            popup = self.popups["current_admission"]
            if hasattr(popup, 'set_entries'):
                filtered, date_info = self._filter_raw_notes_around_admission(raw_notes, 2, 14)
                popup.set_entries(filtered, date_info)
                print(f"[TRIBUNAL] Populated section 8 with {len(filtered)} filtered notes")

        # Section 14: Progress/Mental State - Use notes from last 12 months (from most recent entry)
        from datetime import datetime, timedelta

        def parse_note_date_s14(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # Find the most recent note date
        all_note_dates_s14 = []
        for n in raw_notes:
            note_date = parse_note_date_s14(n.get('date') or n.get('datetime'))
            if note_date:
                all_note_dates_s14.append(note_date)

        if all_note_dates_s14:
            most_recent_date_s14 = max(all_note_dates_s14)
            oldest_date_s14 = min(all_note_dates_s14)
            # 12 months (1 year) before the most recent entry
            one_year_cutoff_s14 = most_recent_date_s14 - timedelta(days=365)

            print(f"[TRIBUNAL] Section 14 DEBUG: Most recent note: {most_recent_date_s14.strftime('%d/%m/%Y')}")
            print(f"[TRIBUNAL] Section 14 DEBUG: Oldest note: {oldest_date_s14.strftime('%d/%m/%Y')}")
            print(f"[TRIBUNAL] Section 14 DEBUG: 1-year cutoff: {one_year_cutoff_s14.strftime('%d/%m/%Y')}")

            # Filter notes from last 12 months (relative to most recent entry)
            notes_with_dates_s14 = []
            for n in raw_notes:
                note_date = parse_note_date_s14(n.get('date') or n.get('datetime'))
                if note_date and note_date >= one_year_cutoff_s14:
                    notes_with_dates_s14.append(n)

            # Sort by date (most recent first)
            recent_progress = sorted(
                notes_with_dates_s14,
                key=lambda x: parse_note_date_s14(x.get('date') or x.get('datetime')),
                reverse=True
            )

            # Show actual date range of filtered notes
            if recent_progress:
                filtered_dates = [parse_note_date_s14(n.get('date') or n.get('datetime')) for n in recent_progress]
                filtered_dates = [d for d in filtered_dates if d]
                if filtered_dates:
                    print(f"[TRIBUNAL] Section 14 DEBUG: Filtered range: {min(filtered_dates).strftime('%d/%m/%Y')} to {max(filtered_dates).strftime('%d/%m/%Y')}")
        else:
            # Fallback: no parseable dates, use first 100 raw notes
            recent_progress = raw_notes[:100]

        print(f"[TRIBUNAL] Section 14: Prepared {len(recent_progress)} notes from last 12 months (total raw: {len(raw_notes)})")

        if "progress" in self.popups:
            popup = self.popups["progress"]
            if hasattr(popup, 'notes'):
                popup.notes = raw_notes

            # Score notes using riskDICT
            self._score_and_display_risk_notes(recent_progress)

            if hasattr(popup, 'set_entries'):
                popup.set_entries(recent_progress, f"{len(recent_progress)} notes (last 12 months)")
                print(f"[TRIBUNAL] Populated section 14 with {len(recent_progress)} notes from last 12 months")

        # Section 17 & 18: Incidents - search all notes using incidentDICT
        from pathlib import Path

        # Load incident terms from dictionary file
        incident_terms = []
        incident_file = Path(__file__).parent / "incidentDICT.txt"
        if incident_file.exists():
            with open(incident_file, 'r', encoding='utf-8') as f:
                for line in f:
                    term = line.strip().lower()
                    if term:
                        incident_terms.append(term)
            print(f"[TRIBUNAL] Loaded {len(incident_terms)} incident terms from incidentDICT.txt")

        exclude_keywords = ['h/o', 'history of', 'previous noted', 'previously noted', 'previous history',
                            'risk of', 'historical risk', 'past risk',
                            # Exclude "no aggression" type phrases
                            'no aggression', 'nil aggression', 'nil aggressive', 'not aggressive',
                            'no aggressive', 'low risk and nil', 'no risk', 'nil on the shift',
                            # Exclude fire safety (not fire incidents)
                            'fire safety', 'fire test', 'gas / fire', 'gas/fire',
                            # Exclude benign mentions
                            'risk and aggression:', 'risk/aggression', 'risk & aggression',
                            # Exclude non-incidents
                            'politely declined', 'hearing 30th', 'managers hearing', 'manangers hearing',
                            'solicitor today', 'no signs of aggression', 'no physical or verbal aggression',
                            'did not display any aggressive', 'no aggressive behaviour',
                            'no challenging or aggressive', 'no episode of aggressive',
                            'no irritable or aggressive', 'no physical/verbal aggression',
                            'no instances of irritability, aggressive',
                            # Exclude negated intoxication
                            'did not appear intoxicated', 'not intoxicated', 'no intoxication',
                            'although did not appear intox', 'did not appear to be intox',
                            # Exclude conditional/future mentions
                            'if no further abusive', 'no further abusive']

        def is_sexual_health_or_history(text):
            """Check if 'sexual' is followed by 'health', 'history', 'hx', or 'transmitted' within 2 words."""
            import re
            # Match 'sexual' followed by 0-2 words then 'health', 'history', 'hx'
            pattern1 = r'\bsexual\s+(?:\w+\s+)?(?:health|history|hx)\b'
            # Match 'sexually transmitted'
            pattern2 = r'\bsexually\s+transmitted\b'
            return bool(re.search(pattern1, text, re.IGNORECASE) or re.search(pattern2, text, re.IGNORECASE))

        # Sort by date helper
        from datetime import datetime, timedelta
        def parse_note_date(date_val):
            if not date_val:
                return None
            if isinstance(date_val, datetime):
                return date_val
            for fmt in ["%Y-%m-%d", "%d/%m/%Y", "%Y-%m-%dT%H:%M:%S", "%d-%m-%Y"]:
                try:
                    return datetime.strptime(str(date_val)[:10], fmt[:min(len(fmt), 10)])
                except:
                    pass
            return None

        # Extract individual incident lines (not full notes)
        import re
        incident_lines = []
        for note in raw_notes:
            content = note.get('content', '') or note.get('text', '') or ''
            # Normalize all line endings to \n
            content = content.replace('\r\n', '\n').replace('\r', '\n')

            date_val = note.get('date') or note.get('datetime')
            date_obj = parse_note_date(date_val)
            date_str = date_obj.strftime('%d/%m/%Y') if date_obj else 'Unknown'

            # Split by newlines first, then by sentences
            all_segments = []
            for line in content.split('\n'):
                # Split by sentence endings (. ! ?) but keep dates like "30/04/2025" intact
                # Split on period followed by space and capital, or end of string
                sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', line)
                all_segments.extend(sentences)

            # Check each segment for incident terms
            for idx, line in enumerate(all_segments):
                line_clean = line.strip()
                # Remove any remaining control characters
                line_clean = ''.join(c for c in line_clean if ord(c) >= 32 or c == '\t')
                line_lower = line_clean.lower()
                if not line_clean or len(line_clean) < 10:
                    continue

                # Skip if contains exclusion keywords
                if any(ex in line_lower for ex in exclude_keywords):
                    continue

                # Skip if "sexual" followed by "health" or "history" (not actual incidents)
                if is_sexual_health_or_history(line_lower):
                    continue

                # Check for incident terms (use word boundary matching for specific problematic terms)
                # "anger" matches inside "manangers", "danger" - causes false positives
                word_boundary_terms = {'anger'}

                for term in incident_terms:
                    if term in word_boundary_terms:
                        # Use word boundaries to avoid "anger" matching "danger"/"manangers"
                        pattern = r'\b' + re.escape(term) + r'\b'
                        if not re.search(pattern, line_lower):
                            continue
                    else:
                        # Most terms use substring match (e.g., "intox" matches "intoxicated")
                        if term not in line_lower:
                            continue

                    # For DATIX or Ulysses, include 2-3 more lines after
                    final_text = line_clean
                    if 'datix' in line_lower or 'ulysses' in line_lower:
                        extra_lines = []
                        for extra_idx in range(1, 4):  # Next 3 lines
                            if idx + extra_idx < len(all_segments):
                                extra_line = all_segments[idx + extra_idx].strip()
                                extra_line = ''.join(c for c in extra_line if ord(c) >= 32 or c == '\t')
                                if extra_line and len(extra_line) > 5:
                                    extra_lines.append(extra_line)
                        if extra_lines:
                            final_text = line_clean + ' | ' + ' | '.join(extra_lines)

                    # Debug: show first 10 matches
                    if len(incident_lines) < 10:
                        print(f"[INCIDENT-DEBUG] Term '{term}' matched in: {line_clean[:80]}...")
                    incident_lines.append({
                        'date': date_str,
                        'date_obj': date_obj,
                        'text': f"{date_str}: {final_text}",
                        'content': f"{date_str}: {final_text}",
                        'term': term
                    })
                    break  # Only one match per line

        # Sort by date (most recent first)
        incident_lines.sort(key=lambda x: x['date_obj'] or datetime.min, reverse=True)

        # Deduplicate by line content
        seen_lines = set()
        incidents = []
        for inc in incident_lines:
            line_key = inc['text'].lower()
            if line_key not in seen_lines:
                seen_lines.add(line_key)
                incidents.append(inc)

        print(f"[TRIBUNAL] Found {len(incidents)} unique incident lines")

        # Use all incidents - no date filtering
        incidents_filtered = incidents
        print(f"[TRIBUNAL] Returning all {len(incidents_filtered)} incidents (no filtering)")

        # Store incident data for later use (when other incident panels are opened)
        self._incident_data = incidents_filtered

        # Section 17: Risk harm - use set_notes to analyze raw notes with risk_overview_panel
        if "risk_harm" in self.popups:
            popup = self.popups["risk_harm"]
            if hasattr(popup, 'set_notes'):
                # TribunalRiskHarmPopup uses analyze_notes_for_risk internally
                popup.set_notes(raw_notes)
                print(f"[TRIBUNAL] Populated section 17 (risk_harm) with {len(raw_notes)} notes for risk analysis")
            elif hasattr(popup, 'set_entries'):
                # Fallback for FixedDataPanel
                label = f"{len(incidents_filtered)} incidents"
                popup.set_entries(incidents_filtered, label)
                print(f"[TRIBUNAL] Populated section 17 with {len(incidents_filtered)} incidents (of {total_incidents} total)")

        # Section 18: Property damage - use set_notes to analyze raw notes with risk_overview_panel
        if "risk_property" in self.popups:
            popup = self.popups["risk_property"]
            if hasattr(popup, 'set_notes'):
                # TribunalRiskPropertyPopup uses analyze_notes_for_risk internally
                popup.set_notes(raw_notes)
                print(f"[TRIBUNAL] Populated section 18 (risk_property) with {len(raw_notes)} notes for risk analysis")
            elif hasattr(popup, 'set_entries'):
                # Fallback for FixedDataPanel
                label = f"{len(incidents_filtered)} incidents"
                popup.set_entries(incidents_filtered, label)
                print(f"[TRIBUNAL] Populated section 18 with {len(incidents_filtered)} incidents")

        # Section 5: Forensic History - populate forensic history panel with notes analysis + extracted data
        forensic_items = get_category_items("Forensic History")
        if not forensic_items:
            # Try alternative category names
            forensic_items = get_category_items("forensic history")
        if not forensic_items:
            forensic_items = get_category_items("FORENSIC")

        # Get raw notes for forensic risk analysis
        raw_notes = getattr(self, '_extracted_raw_notes', [])

        # Skip forensic notes population if report data was imported for forensic
        if "forensic" in self.popups and not (hasattr(self, '_imported_report_data') and 'forensic' in self._imported_report_data):
            popup = self.popups["forensic"]
            if hasattr(popup, 'set_forensic_data'):
                # Use combined notes analysis + extracted data
                popup.set_forensic_data(raw_notes, forensic_items)
                print(f"[TRIBUNAL] Populated section 5 forensic panel with notes analysis + {len(forensic_items) if forensic_items else 0} extracted entries")
            elif hasattr(popup, 'set_entries') and forensic_items:
                popup.set_entries(forensic_items, f"{len(forensic_items)} entries")
                print(f"[TRIBUNAL] Populated section 5 forensic panel with {len(forensic_items)} entries")

        # Store forensic items and notes for later use (when popup is created)
        if not hasattr(self, '_pending_forensic_data'):
            self._pending_forensic_data = {}
        self._pending_forensic_data['forensic'] = forensic_items or []
        self._pending_forensic_data['forensic_notes'] = raw_notes
        print(f"[TRIBUNAL] Stored {len(forensic_items) if forensic_items else 0} forensic entries + {len(raw_notes)} notes for forensic popup")

        # Pre-fill medications in section 12 (treatment) - same as nursing section 4
        if "treatment" in self.popups:
            self._prefill_medications_from_notes()

        print("[TRIBUNAL] Fixed panels populated with extracted data")

    def _go_back(self):
        """Return to reports selection page."""
        self.go_back.emit()

    def _export_docx(self):
        """Export report to DOCX using T131 template with table-based input boxes."""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from docx import Document
        from docx.shared import Pt
        from datetime import datetime
        import os
        import shutil
        import re

        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Export Psychiatric Tribunal Report",
            f"tribunal_report_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            # Find template path - use new template with real input boxes
            template_path = resource_path("templates", "t131_template_new.docx")

            if not os.path.exists(template_path):
                QMessageBox.critical(self, "Template Error", f"T131 template not found at:\n{template_path}")
                return

            # Copy template to destination
            shutil.copy(template_path, file_path)

            # Open the copied template
            doc = Document(file_path)

            # Helper to get card content (or popup state for heading-only sections)
            def get_content(key):
                # For heading-only sections, get content from popup's generate_text()
                if key in self.HEADING_ONLY_SECTIONS:
                    if key in self.popups and hasattr(self.popups[key], 'generate_text'):
                        return self.popups[key].generate_text()
                    return ""
                # For regular cards, get from editor
                if key in self.cards:
                    return self.cards[key].editor.toPlainText().strip()
                return ""

            # Helper to extract detail text after Yes/No prefix
            def extract_detail(content):
                if not content:
                    return ""
                detail = content
                lower = detail.lower()
                for prefix in ["yes, ", "yes - ", "yes-", "yes,", "yes.", "no, ", "no - ", "no-", "no,", "no."]:
                    if lower.startswith(prefix):
                        detail = detail[len(prefix):].strip()
                        break
                if lower == "yes" or lower == "no":
                    return ""
                return detail

            # Helper to check if content indicates Yes answer
            def has_yes_content(content):
                """Returns True if content indicates Yes answer."""
                if not content:
                    return False
                lower = content.lower().strip()
                no_patterns = [
                    "no ", "no,", "no-", "no.",
                    "there have been no", "there is no",
                    "no incidents", "no factors", "no adjustments",
                    "does not have", "not applicable", "n/a",
                ]
                if lower == "no":
                    return False
                for pattern in no_patterns:
                    if lower.startswith(pattern) or pattern in lower:
                        return False
                return True

            # Helper to set table cell text
            def set_cell(table_idx, row, col, text):
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    if row < len(table.rows) and col < len(table.rows[row].cells):
                        cell = table.cell(row, col)
                        cell.text = text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(11)

            # ============================================================
            # PARSE PATIENT DETAILS
            # ============================================================
            pd_content = get_content("patient_details")
            pd_lines = pd_content.split('\n') if pd_content else []
            full_name = dob = residence = ""

            # Track which lines contain specific info
            address_lines = []
            collecting_address = False

            for i, line in enumerate(pd_lines):
                line_lower = line.lower().strip()

                if "name:" in line_lower or "full name:" in line_lower:
                    full_name = line.split(":", 1)[-1].strip()
                    collecting_address = False
                elif "date of birth:" in line_lower or "dob:" in line_lower:
                    dob = line.split(":", 1)[-1].strip()
                    collecting_address = False
                elif any(x in line_lower for x in ["residence:", "address:", "usual place"]):
                    # Start collecting address (may span multiple lines)
                    addr_part = line.split(":", 1)[-1].strip()
                    if addr_part:
                        address_lines.append(addr_part)
                    collecting_address = True
                elif collecting_address and line.strip():
                    # Continue collecting address lines until we hit another field
                    if not any(x in line_lower for x in ["name:", "dob:", "gender:", "nhs", "hospital"]):
                        address_lines.append(line.strip())
                    else:
                        collecting_address = False

            # Combine address lines
            if address_lines:
                residence = ", ".join(address_lines)

            # If no structured name, use first non-empty line
            if not full_name and pd_lines:
                for line in pd_lines:
                    if line.strip() and not any(x in line.lower() for x in ["dob", "date of birth", "residence", "address"]):
                        full_name = line.strip()
                        break

            print(f"[EXPORT] Patient details - Name: {full_name}, DOB: {dob}, Address: {residence}")

            # Table 1: Full name (1x1)
            set_cell(1, 0, 0, full_name)

            # Table 2: DOB character boxes (1x8)
            if dob:
                dob_clean = ""
                date_formats = [
                    ("%d/%m/%Y", None), ("%d-%m-%Y", None), ("%d.%m.%Y", None),
                    ("%d %B %Y", None), ("%d %b %Y", None), ("%Y-%m-%d", None),
                ]
                parsed_date = None
                for fmt, _ in date_formats:
                    try:
                        parsed_date = datetime.strptime(dob.strip(), fmt)
                        break
                    except:
                        continue

                if parsed_date:
                    dob_clean = parsed_date.strftime("%d%m%Y")
                else:
                    dob_clean = re.sub(r'[^0-9]', '', dob)
                    if len(dob_clean) == 6:
                        year_part = dob_clean[4:6]
                        year = "20" + year_part if int(year_part) < 50 else "19" + year_part
                        dob_clean = dob_clean[:4] + year

                if len(doc.tables) > 2:
                    dob_table = doc.tables[2]
                    for i, char in enumerate(dob_clean[:8]):
                        if i < len(dob_table.columns):
                            dob_table.cell(0, i).text = char

            # Table 3: Residence (1x1)
            set_cell(3, 0, 0, residence)

            # ============================================================
            # PARSE SIGNATURE INFO FOR RC NAME
            # ============================================================
            sig_content = get_content("signature")
            sig_lines = sig_content.split('\n') if sig_content else []
            sig_name = sig_role = ""
            sig_date = datetime.now().strftime("%d/%m/%Y")

            for line in sig_lines:
                line_lower = line.lower()
                if "signed:" in line_lower or "name:" in line_lower:
                    sig_name = line.split(":", 1)[-1].strip()
                elif "designation:" in line_lower or "role:" in line_lower:
                    sig_role = line.split(":", 1)[-1].strip()
                elif "date:" in line_lower:
                    sig_date = line.split(":", 1)[-1].strip()

            if not sig_name and sig_lines:
                for line in sig_lines:
                    if line.strip() and not any(x in line.lower() for x in ["date", "designation", "role", "registration"]):
                        sig_name = line.strip()
                        break

            # ============================================================
            # PARSE AUTHOR INFO (may have structured content)
            # ============================================================
            author_content = get_content("author")
            author_name = ""
            author_role = ""
            author_date = ""

            if author_content:
                author_lines = author_content.split('\n')
                for line in author_lines:
                    line_lower = line.lower()
                    if "name" in line_lower and ":" in line:
                        author_name = line.split(":", 1)[-1].strip()
                    elif "role:" in line_lower or "designation:" in line_lower:
                        author_role = line.split(":", 1)[-1].strip()
                    elif "date:" in line_lower:
                        author_date = line.split(":", 1)[-1].strip()

                # If no structured name found, use first non-empty line
                if not author_name:
                    for line in author_lines:
                        stripped = line.strip()
                        if stripped and ":" not in stripped:
                            author_name = stripped
                            break

            # Use parsed values, falling back to signature info
            rc_name = author_name or sig_name
            rc_role = author_role or sig_role
            rc_date = author_date or sig_date

            # Table 4: RC Name (1x1)
            set_cell(4, 0, 0, rc_name)

            # ============================================================
            # NESTED TABLE IN TABLE 0 - Hospital info box (top right)
            # Structure: Row 0=Hospital logo title, Row 1=Your name title,
            #            Row 2=name entry, Row 3=Your role title,
            #            Row 4=role entry, Row 5=Date title, Row 6=date entry
            # ============================================================
            if len(doc.tables) > 0:
                table0 = doc.tables[0]
                if len(table0.rows[0].cells) > 1:
                    cell_with_nested = table0.rows[0].cells[1]
                    if cell_with_nested.tables:
                        nested_table = cell_with_nested.tables[0]
                        # Row 2: Your name entry
                        if len(nested_table.rows) > 2:
                            nested_table.rows[2].cells[0].text = rc_name
                        # Row 4: Your role entry
                        if len(nested_table.rows) > 4:
                            nested_table.rows[4].cells[0].text = rc_role
                        # Row 6: Date entry
                        if len(nested_table.rows) > 6:
                            nested_table.rows[6].cells[0].text = rc_date

            # ============================================================
            # SECTION INPUT BOXES (Tables 5-21)
            # T131 Table mapping: table_idx -> (card_key, use_extract_detail)
            # Note: learning_disability, detention_required, s2_detention,
            #       other_detention are checkbox-only (no content tables)
            # ============================================================
            table_mapping = {
                5: ("factors_hearing", True),       # Are there any factors...
                6: ("adjustments", True),           # Are there any adjustments...
                7: ("forensic", False),             # Index offence and forensic history
                8: ("previous_mh_dates", False),    # Dates of previous MH involvement
                9: ("previous_admission_reasons", False),  # Reasons for previous admission
                10: ("current_admission", False),   # Circumstances of current admission
                11: ("diagnosis", True),            # Mental disorder/diagnosis
                12: ("treatment", False),           # Medical treatment
                13: ("strengths", False),           # Strengths/positive factors
                14: ("progress", False),            # Progress, behaviour, capacity
                15: ("compliance", False),          # Compliance with treatment
                16: ("mca_dol", False),             # MCA/DoL consideration
                17: ("risk_harm", False),           # Harm to self/others
                18: ("risk_property", False),       # Property damage
                19: ("discharge_risk", True),       # After s2/other/discharge checkboxes
                20: ("community", False),           # Community risk management
                21: ("recommendations", True),      # Recommendations to tribunal
                # Table 22: Signature box
                # Table 23: Date character boxes
            }

            # Sections with Yes/No checkboxes - only fill table if Yes
            yes_no_sections = {"factors_hearing", "adjustments", "diagnosis",
                               "discharge_risk", "recommendations"}

            for table_idx, (card_key, use_extract) in table_mapping.items():
                content = get_content(card_key)
                if content:
                    # For Yes/No sections, only fill table if answer is Yes
                    if card_key in yes_no_sections:
                        if not has_yes_content(content):
                            continue  # Skip table entry for No answers
                    if use_extract:
                        content = extract_detail(content) or content
                    set_cell(table_idx, 0, 0, content)

            # ============================================================
            # CHECKBOX SECTIONS - mark ☐ → ☒ based on content
            # Detect sections by question text patterns, not numbers
            # ============================================================
            checkbox_patterns = {
                "factors that may affect": "factors_hearing",
                "any adjustments": "adjustments",
                "suffering from a mental disorder": "diagnosis",
                "learning disability": "learning_disability",
                "requires the patient to be detained": "detention_required",
                "section 2 cases": "s2_detention",
                "in all other cases": "other_detention",
                "discharged from hospital": "discharge_risk",
                "recommendations to the tribunal": "recommendations",
            }

            current_section_key = None
            for para in doc.paragraphs:
                text = para.text.strip()
                lower_text = text.lower()

                # Detect which section we're in by matching question patterns
                for pattern, card_key in checkbox_patterns.items():
                    if pattern in lower_text:
                        current_section_key = card_key
                        break

                # Mark checkboxes based on content
                if current_section_key and "☐" in para.text:
                    content = get_content(current_section_key)
                    is_yes = has_yes_content(content)

                    is_yes_line = "yes" in lower_text and "no" not in lower_text[:10]
                    is_no_line = lower_text.strip().startswith("☐ no") or (lower_text.strip().startswith("☐") and "no" in lower_text and "yes" not in lower_text)

                    if is_yes_line and is_yes:
                        for run in para.runs:
                            if "☐" in run.text:
                                run.text = run.text.replace("☐", "☒", 1)
                                break
                    elif is_no_line and not is_yes and content:
                        for run in para.runs:
                            if "☐" in run.text:
                                run.text = run.text.replace("☐", "☒", 1)
                                break

            # ============================================================
            # SIGNATURE - Table 22 is the signature box
            # ============================================================
            from docx.shared import Inches

            sig_text_parts = []
            if sig_name:
                sig_text_parts.append(sig_name)
            if sig_role:
                sig_text_parts.append(sig_role)

            for line in sig_lines:
                line_lower = line.lower()
                if "qualification" in line_lower:
                    sig_text_parts.append(line.split(":", 1)[-1].strip())
                elif "registration:" in line_lower or "gmc:" in line_lower:
                    reg = line.split(":", 1)[-1].strip()
                    sig_text_parts.append(f"Registration: {reg}")

            # Put signature image and text in Table 22
            if 22 < len(doc.tables):
                sig_cell = doc.tables[22].cell(0, 0)
                sig_cell.text = ""  # Clear existing content

                # Add signature image if exists
                sig_image_path = os.path.expanduser("~/MyPsychAdmin/signature.png")
                if os.path.exists(sig_image_path):
                    sig_para = sig_cell.paragraphs[0]
                    sig_run = sig_para.add_run()
                    sig_run.add_picture(sig_image_path, width=Inches(1.5))
                    # Add new paragraph for text
                    sig_cell.add_paragraph("\n".join(sig_text_parts))
                else:
                    sig_cell.text = "\n".join(sig_text_parts)

            # ============================================================
            # TABLE 23: Date character boxes (1x8) - ddmmyyyy format
            # ============================================================
            if sig_date and len(doc.tables) > 23:
                date_table = doc.tables[23]
                date_clean = ""

                date_formats = [
                    ("%d/%m/%Y", None), ("%d-%m-%Y", None), ("%d.%m.%Y", None),
                    ("%d %B %Y", None), ("%d %b %Y", None), ("%Y-%m-%d", None),
                    ("%B %d, %Y", None), ("%b %d, %Y", None),
                ]

                parsed_date = None
                for fmt, _ in date_formats:
                    try:
                        parsed_date = datetime.strptime(sig_date.strip(), fmt)
                        break
                    except:
                        continue

                if parsed_date:
                    date_clean = parsed_date.strftime("%d%m%Y")
                else:
                    date_clean = re.sub(r'[^0-9]', '', sig_date)
                    if len(date_clean) == 6:
                        year_part = date_clean[4:6]
                        year = "20" + year_part if int(year_part) < 50 else "19" + year_part
                        date_clean = date_clean[:4] + year

                for i, char in enumerate(date_clean[:8]):
                    if i < len(date_table.columns):
                        date_table.cell(0, i).text = char

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Report exported to:\n{file_path}")

        except Exception as e:
            import traceback
            traceback.print_exc()
            QMessageBox.critical(self, "Export Error", f"Failed to export report:\n{str(e)}")

    def get_report_data(self) -> dict:
        """Get all report data as a dictionary."""
        data = {}
        for title, key in self.SECTIONS:
            # For heading-only sections, get from popup
            if key in self.HEADING_ONLY_SECTIONS:
                if key in self.popups and hasattr(self.popups[key], 'generate_text'):
                    data[key] = self.popups[key].generate_text()
                else:
                    data[key] = ""
            elif key in self.cards:
                data[key] = self.cards[key].editor.toPlainText().strip()
        return data

    def set_report_data(self, data: dict):
        """Set report data from a dictionary."""
        for key, content in data.items():
            # For heading-only sections, restore popup state
            if key in self.HEADING_ONLY_SECTIONS:
                if key in self.popups and hasattr(self.popups[key], 'restore_state'):
                    content_lower = content.lower() if content else ""
                    is_yes = content_lower.startswith("yes")

                    # LearningDisabilityPopup has different state format
                    if key == "learning_disability":
                        has_aggressive = "aggressive" in content_lower and "is associated" in content_lower
                        self.popups[key].restore_state({
                            "has_ld": "yes" if is_yes else "no",
                            "aggressive_conduct": "yes" if has_aggressive else "no"
                        })
                    else:
                        # SimpleYesNoPopup format
                        self.popups[key].restore_state({"answer": "yes" if is_yes else "no"})
            elif key in self.cards:
                self.cards[key].editor.setPlainText(content)
