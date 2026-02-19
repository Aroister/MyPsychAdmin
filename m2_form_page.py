# ================================================================
#  M2 FORM PAGE — Report Barring Discharge by Nearest Relative
#  Mental Health Act 1983 - Form M2 Regulation 25(1)(a) and (b)
#  Section 25 — Report barring discharge by nearest relative
#  CARD/POPUP LAYOUT VERSION
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime, QEvent
from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QPushButton, QFileDialog, QMessageBox, QToolButton,
    QRadioButton, QButtonGroup, QSpinBox, QComboBox, QCheckBox,
    QCompleter, QStyleFactory, QSlider, QSizePolicy, QStackedWidget,
    QSplitter
)
from PySide6.QtGui import QColor
from PySide6.QtWidgets import QGraphicsDropShadowEffect
from background_history_popup import ResizableSection
from shared_widgets import create_zoom_row
from utils.resource_path import resource_path
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []


# ================================================================
# NO-WHEEL WIDGETS (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelSpinBox(QSpinBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelTimeEdit(QTimeEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# CARD WIDGET FOR M2 FORM
# ================================================================
class M2CardWidget(QFrame):
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.title = title
        self.key = key
        self._active = False

        self.setObjectName("m2Card")
        self.setStyleSheet("""
            QFrame#m2Card {
                background: white;
                border-radius: 12px;
                border: 1px solid #e5e7eb;
            }
            QFrame#m2Card:hover {
                border-color: #be185d;
                background: #fdf2f8;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QLabel#cardTitle {
                font-size: 20px;
                font-weight: 600;
                color: #be185d;
                padding-bottom: 4px;
            }
            QFrame#divider {
                background: rgba(0,0,0,0.10);
                height: 1px;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 12)
        layout.setSpacing(0)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setObjectName("cardTitle")
        self.title_label.setFixedHeight(24)
        self.title_label.setCursor(Qt.CursorShape.PointingHandCursor)
        self.title_label.mousePressEvent = lambda e: self.clicked.emit(self.key)
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        divider = QFrame()
        divider.setObjectName("divider")
        divider.setFixedHeight(1)
        layout.addWidget(divider)

        # Scrollable content area
        self.content_scroll = QScrollArea()
        self.content_scroll.setWidgetResizable(True)
        self.content_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        self.content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        content_widget = QWidget()
        content_widget.setStyleSheet("background: transparent;")
        content_layout = QVBoxLayout(content_widget)
        content_layout.setContentsMargins(0, 8, 0, 4)
        content_layout.setSpacing(0)

        self.content = QLabel("")
        self.content.setWordWrap(True)
        self.content.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.content.setStyleSheet("""
            QLabel {
                font-size: 16px;
                color: #374151;
                padding: 4px;
                background: transparent;
            }
        """)
        content_layout.addWidget(self.content)
        content_layout.addStretch()

        self.content_scroll.setWidget(content_widget)
        layout.addWidget(self.content_scroll, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=16)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

        self.setCursor(Qt.CursorShape.PointingHandCursor)

    def set_active(self, active: bool):
        if active:
            self.setStyleSheet("""
                QFrame#m2Card {
                    background: #fdf2f8;
                    border: 2px solid #be185d;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
                QLabel#cardTitle {
                    font-size: 20px;
                    font-weight: 600;
                    color: #be185d;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#m2Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#m2Card:hover {
                    border-color: #be185d;
                    background: #fdf2f8;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
                QLabel#cardTitle {
                    font-size: 20px;
                    font-weight: 600;
                    color: #be185d;
                }
            """)

    def set_content(self, text: str):
        """Set the summary content displayed on the card."""
        self.content.setText(text)

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)


# ================================================================
# M2 FORM PAGE - MAIN CLASS
# ================================================================
class M2FormPage(QWidget):
    """Page for completing MHA Form M2 - Report Barring Discharge."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}
        self._active_section = None
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return dict(details)

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his", "self": "himself"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her", "self": "herself"}
        else:
            return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their", "self": "themselves"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #be185d; border-bottom: 1px solid rgba(0,0,0,0.1);")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form M2 — Report Barring Discharge by Nearest Relative")
        title.setStyleSheet("font-size: 18px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        main_layout.addWidget(self.toolbar)

        # ==================================================
        # MAIN SPLIT AREA (Cards on left, Popup on right)
        # ==================================================
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QHBoxLayout(content)
        content_layout.setContentsMargins(12, 12, 12, 12)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #be185d, stop:0.7 #be185d, stop:1 transparent);
            }
        """)

        # Sections for cards (consolidated - Reasons includes diagnosis, legal, risk, informal)
        self.sections = [
            ("Hospital & Notice Details", "hospital"),
            ("Patient Details", "patient"),
            ("Reasons", "reasons"),
            ("RC & Signature", "signature"),
        ]

        # ---------------- LEFT: CARDS ----------------
        self.cards_holder = QScrollArea()
        self.cards_holder.setWidgetResizable(True)
        self.cards_holder.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.cards_holder.setFrameShape(QScrollArea.Shape.NoFrame)
        self.cards_holder.setStyleSheet("QScrollArea { background: transparent; border: none; }")
        self.main_splitter.addWidget(self.cards_holder)

        cards_container = QWidget()
        cards_container.setStyleSheet("background: #f9fafb;")
        cards_layout = QVBoxLayout(cards_container)
        cards_layout.setContentsMargins(30, 30, 30, 30)
        cards_layout.setSpacing(20)

        self.card_sections = {}
        for title, key in self.sections:
            # Create resizable section for each card
            section = ResizableSection()
            section.set_content_height(160)
            section._min_height = 100
            section._max_height = 350

            card = M2CardWidget(title, key)
            card.clicked.connect(self._activate_section)
            self.cards[key] = card
            section.set_content(card)
            self.card_sections[key] = section
            cards_layout.addWidget(section)

        cards_layout.addStretch()
        self.cards_holder.setWidget(cards_container)

        # ---------------- RIGHT: POPUP PANEL ----------------
        self.popup_panel = QFrame()
        self.popup_panel.setStyleSheet("""
            QFrame {
                background: white;
                border: none;
                border-radius: 10px;
            }
        """)
        self.main_splitter.addWidget(self.popup_panel)
        self.main_splitter.setSizes([280, 700])
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        panel_layout = QVBoxLayout(self.popup_panel)
        panel_layout.setContentsMargins(20, 20, 20, 20)
        panel_layout.setSpacing(12)

        self.panel_title = QLabel("Select a section")
        self.panel_title.setStyleSheet("""
            font-size: 18px;
            font-weight: 700;
            color: #be185d;
            background: rgba(253,242,248,0.85);
            padding: 8px 12px;
            border-radius: 8px;
        """)
        panel_layout.addWidget(self.panel_title)

        # Popup stack
        self.popup_stack = QStackedWidget()
        panel_layout.addWidget(self.popup_stack, 1)

        # Build all popup panels
        self._build_popup_panels()

        # Default to first section
        self._activate_section("hospital")

    def _activate_section(self, key: str):
        """Activate a section card and show its popup."""
        self._active_section = key

        # Update card highlights
        for k, card in self.cards.items():
            card.set_active(k == key)

        # Show appropriate popup
        index = [s[1] for s in self.sections].index(key)
        self.popup_stack.setCurrentIndex(index)

        # Update panel title
        title = [s[0] for s in self.sections if s[1] == key][0]
        self.panel_title.setText(title)

    def _build_popup_panels(self):
        """Build popup panels for each section."""
        self._build_hospital_popup()
        self._build_patient_popup()
        self._build_reasons_popup()  # Combined: diagnosis, legal, risk, informal + preview
        self._build_signature_popup()

        # Initialize cards with default date values
        self._update_hospital_card()
        self._update_signature_card()

    def _create_popup_scroll(self) -> tuple:
        """Create a scrollable popup container."""
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setStyleSheet("QScrollArea { background: transparent; }")

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        layout = QVBoxLayout(container)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(12)

        scroll.setWidget(container)
        return scroll, layout

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 17px;
            }
            QLineEdit:focus { border-color: #be185d; }
        """)
        return edit

    def _create_date_edit(self) -> NoWheelDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 17px;
            }
        """)
        return date_edit

    def _create_time_edit(self) -> NoWheelTimeEdit:
        time_edit = NoWheelTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("""
            QTimeEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 17px;
            }
        """)
        return time_edit

    def _create_section_label(self, text: str) -> QLabel:
        lbl = QLabel(text)
        lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151;")
        return lbl

    # ================================================================
    # POPUP BUILDERS
    # ================================================================

    def _build_hospital_popup(self):
        """Build Hospital & Notice Details popup."""
        scroll, layout = self._create_popup_scroll()

        layout.addWidget(self._create_section_label("To the managers of:"))
        self.hospital = self._create_line_edit("Hospital name and address")
        layout.addWidget(self.hospital)

        layout.addWidget(self._create_section_label("Notice of intention to discharge given by:"))
        self.nearest_relative = self._create_line_edit("Nearest relative name")
        layout.addWidget(self.nearest_relative)

        time_date_row = QHBoxLayout()
        time_date_row.setSpacing(12)

        time_lbl = QLabel("At time:")
        time_lbl.setStyleSheet("font-size: 16px; color: #374151;")
        time_date_row.addWidget(time_lbl)
        self.notice_time = self._create_time_edit()
        self.notice_time.setFixedWidth(100)
        time_date_row.addWidget(self.notice_time)

        date_lbl = QLabel("On date:")
        date_lbl.setStyleSheet("font-size: 16px; color: #374151;")
        time_date_row.addWidget(date_lbl)
        self.notice_date = self._create_date_edit()
        self.notice_date.setFixedWidth(140)
        time_date_row.addWidget(self.notice_date)

        time_date_row.addStretch()
        layout.addLayout(time_date_row)

        layout.addStretch()

        # Auto-sync to card
        self.hospital.textChanged.connect(self._update_hospital_card)
        self.nearest_relative.textChanged.connect(self._update_hospital_card)
        self.notice_time.timeChanged.connect(self._update_hospital_card)
        self.notice_date.dateChanged.connect(self._update_hospital_card)

        self.popup_stack.addWidget(scroll)

    def _update_hospital_card(self):
        parts = []
        if self.hospital.text().strip():
            parts.append(self.hospital.text().strip())
        if self.nearest_relative.text().strip():
            parts.append(self.nearest_relative.text().strip())
        self.cards["hospital"].set_content("\n".join(parts) if parts else "Click to enter details")

    def _build_patient_popup(self):
        """Build Patient Details popup."""
        scroll, layout = self._create_popup_scroll()

        layout.addWidget(self._create_section_label("Intention to discharge:"))
        self.patient_name = self._create_line_edit("Patient full name")
        layout.addWidget(self.patient_name)

        demo_frame = QFrame()
        demo_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 8px; }")
        demo_layout = QVBoxLayout(demo_frame)
        demo_layout.setContentsMargins(12, 10, 12, 10)
        demo_layout.setSpacing(8)

        demo_header = QLabel("Demographics")
        demo_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #166534;")
        demo_layout.addWidget(demo_header)

        # Row 1: Age and Gender
        demo_row1 = QHBoxLayout()
        demo_row1.setSpacing(8)

        # Age - visible for clinical reasons text generation
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 16px; color: #374151;")
        demo_row1.addWidget(age_lbl)
        self.age_spin = NoWheelSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setFixedWidth(55)
        self.age_spin.setStyleSheet("font-size: 16px; padding: 4px;")
        demo_row1.addWidget(self.age_spin)

        # Gender
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("Male")
        self.gender_female = QRadioButton("Female")
        self.gender_other = QRadioButton("Other")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 16px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
            self.gender_group.addButton(rb)
            demo_row1.addWidget(rb)
        demo_row1.addStretch()
        demo_layout.addLayout(demo_row1)

        # Row 2: Ethnicity (under Age)
        demo_row2 = QHBoxLayout()
        demo_row2.setSpacing(8)
        eth_lbl = QLabel("Ethnicity:")
        eth_lbl.setStyleSheet("font-size: 16px; color: #374151;")
        demo_row2.addWidget(eth_lbl)
        self.ethnicity_combo = NoWheelComboBox()
        self.ethnicity_combo.addItem("Not specified")
        self.ethnicity_combo.addItems([
            "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race"
        ])
        self.ethnicity_combo.setFixedWidth(140)
        self.ethnicity_combo.setStyleSheet("font-size: 15px; padding: 4px;")
        demo_row2.addWidget(self.ethnicity_combo)
        demo_row2.addStretch()
        demo_layout.addLayout(demo_row2)

        layout.addWidget(demo_frame)

        layout.addStretch()

        # Auto-sync to card (age_spin and ethnicity_combo not shown on card, only gender)
        self.patient_name.textChanged.connect(self._update_patient_card)
        self.gender_male.toggled.connect(self._update_patient_card)
        self.gender_female.toggled.connect(self._update_patient_card)
        self.gender_other.toggled.connect(self._update_patient_card)

        # Connect demographics to reasons text generation for clinical preview
        self.patient_name.textChanged.connect(self._generate_reasons_text)
        self.age_spin.valueChanged.connect(self._generate_reasons_text)
        self.gender_male.toggled.connect(self._generate_reasons_text)
        self.gender_female.toggled.connect(self._generate_reasons_text)
        self.gender_other.toggled.connect(self._generate_reasons_text)
        self.ethnicity_combo.currentIndexChanged.connect(self._generate_reasons_text)

        self.popup_stack.addWidget(scroll)

    def _update_patient_card(self):
        parts = []
        if self.patient_name.text().strip():
            parts.append(self.patient_name.text().strip())
        demo_parts = []
        # Age and ethnicity hidden from card display (kept for backend clinical text)
        if self.gender_male.isChecked():
            demo_parts.append("Male")
        elif self.gender_female.isChecked():
            demo_parts.append("Female")
        if demo_parts:
            parts.append(", ".join(demo_parts))
        self.cards["patient"].set_content("\n".join(parts) if parts else "Click to enter patient details")

    def _build_reasons_popup(self):
        """Build combined Reasons popup with FIXED preview at top, scrollable inputs below."""
        # Main container for the whole popup (not scrollable)
        main_container = QWidget()
        main_container.setStyleSheet("background: transparent;")
        main_layout = QVBoxLayout(main_container)
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(8)

        # ============================================================
        # FIXED HEADER: "I am of the opinion..."
        # ============================================================
        op_lbl = QLabel("I am of the opinion that the patient, if discharged, would be likely to act in a manner dangerous to other persons or to himself/herself.")
        op_lbl.setWordWrap(True)
        op_lbl.setStyleSheet("font-size: 16px; color: #374151; padding: 8px; background: #fef2f2; border-radius: 6px;")
        main_layout.addWidget(op_lbl)

        # Hidden label for state storage (preview removed - card auto-syncs)
        self.reasons = QLabel("")
        self.reasons.hide()

        # ============================================================
        # SCROLLABLE INPUTS SECTION
        # ============================================================
        inputs_scroll = QScrollArea()
        inputs_scroll.setWidgetResizable(True)
        inputs_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        inputs_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        inputs_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        inputs_scroll.setStyleSheet("QScrollArea { background: transparent; }")

        inputs_container = QWidget()
        inputs_container.setStyleSheet("background: transparent;")
        inputs_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        layout = QVBoxLayout(inputs_container)
        layout.setContentsMargins(0, 8, 0, 8)
        layout.setSpacing(12)

        # ============================================================
        # SECTION 1: MENTAL DISORDER (ICD-10)
        # ============================================================
        md_frame = QFrame()
        md_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 8px; }")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(8)

        md_header = QLabel("Mental Disorder (ICD-10)")
        md_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        # Primary diagnosis dropdown with grouped items
        self.dx_primary = NoWheelComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select primary diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_primary.currentTextChanged.connect(self._generate_reasons_text)
        md_layout.addWidget(self.dx_primary)

        # Secondary diagnosis dropdown
        self.dx_secondary = NoWheelComboBox()
        self.dx_secondary.setEditable(True)
        self.dx_secondary.lineEdit().setPlaceholderText("Secondary diagnosis (optional)...")
        self.dx_secondary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_secondary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_secondary.addItem(f"── {group_name} ──", None)
            idx = self.dx_secondary.count() - 1
            self.dx_secondary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_secondary.addItem(dx, dx)
        completer2 = QCompleter(ICD10_FLAT, self.dx_secondary)
        completer2.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer2.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_secondary.setCompleter(completer2)
        self.dx_secondary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        self.dx_secondary.currentTextChanged.connect(self._generate_reasons_text)
        md_layout.addWidget(self.dx_secondary)

        layout.addWidget(md_frame)

        # ============================================================
        # SECTION 2: LEGAL CRITERIA (Nature/Degree)
        # ============================================================
        lc_frame = QFrame()
        lc_frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 8px; }")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(6)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.relapsing_cb.toggled.connect(self._generate_reasons_text)
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.treatment_resistant_cb.toggled.connect(self._generate_reasons_text)
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.chronic_cb.toggled.connect(self._generate_reasons_text)
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(120)
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 16px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.degree_details.textChanged.connect(self._generate_reasons_text)
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        layout.addWidget(lc_frame)

        # ============================================================
        # SECTION 3: RISK ASSESSMENT (Health/Safety)
        # ============================================================
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 17px; font-weight: 700; color: #374151;")
        layout.addWidget(nec_lbl)

        # Health frame
        health_frame = QFrame()
        health_frame.setStyleSheet("QFrame { background: #fef3c7; border: none; border-radius: 8px; }")
        health_layout = QVBoxLayout(health_frame)
        health_layout.setContentsMargins(12, 10, 12, 10)
        health_layout.setSpacing(4)

        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        health_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.poor_compliance_cb.toggled.connect(self._generate_reasons_text)
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
        self.limited_insight_cb.toggled.connect(self._generate_reasons_text)
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.textChanged.connect(self._generate_reasons_text)
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        health_layout.addWidget(self.health_options)

        layout.addWidget(health_frame)

        # Safety frame
        safety_frame = QFrame()
        safety_frame.setStyleSheet("QFrame { background: #fef2f2; border: none; border-radius: 8px; }")
        safety_layout = QVBoxLayout(safety_frame)
        safety_layout.setContentsMargins(12, 10, 12, 10)
        safety_layout.setSpacing(4)

        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        safety_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # Self section
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 15px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.self_hist_neglect.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.self_hist_risky.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.self_hist_harm.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_hist_harm)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 15px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.self_curr_neglect.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.self_curr_risky.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.self_curr_harm.toggled.connect(self._generate_reasons_text)
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others section
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 15px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence")
        self.others_hist_violence.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_hist_violence.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_hist_verbal.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_verbal)

        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_sexual.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_hist_sexual.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_sexual)

        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_stalking.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_hist_stalking.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_stalking)

        self.others_hist_arson = QCheckBox("Arson")
        self.others_hist_arson.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_hist_arson.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_hist_arson)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 15px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence")
        self.others_curr_violence.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_curr_violence.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_curr_verbal.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_sexual.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_curr_sexual.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_sexual)

        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_stalking.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_curr_stalking.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_stalking)

        self.others_curr_arson = QCheckBox("Arson")
        self.others_curr_arson.setStyleSheet("font-size: 15px; color: #9ca3af;")
        self.others_curr_arson.toggled.connect(self._generate_reasons_text)
        others_opt_layout.addWidget(self.others_curr_arson)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        safety_layout.addWidget(self.safety_options)

        layout.addWidget(safety_frame)

        # ============================================================
        # SECTION 4: INFORMAL NOT INDICATED
        # ============================================================
        inf_frame = QFrame()
        inf_frame.setStyleSheet("QFrame { background: #fef2f2; border: none; border-radius: 8px; }")
        inf_layout = QVBoxLayout(inf_frame)
        inf_layout.setContentsMargins(12, 10, 12, 10)
        inf_layout.setSpacing(6)

        inf_header = QLabel("Informal Not Indicated Because:")
        inf_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #991b1b;")
        inf_layout.addWidget(inf_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed previously")
        self.tried_failed_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.tried_failed_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.tried_failed_cb)

        self.insight_cb = QCheckBox("Lack of Insight")
        self.insight_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.insight_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.insight_cb)

        self.compliance_cb = QCheckBox("Compliance Issues")
        self.compliance_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.compliance_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.compliance_cb)

        self.supervision_cb = QCheckBox("Needs MHA Supervision")
        self.supervision_cb.setStyleSheet("font-size: 16px; color: #374151;")
        self.supervision_cb.toggled.connect(self._generate_reasons_text)
        inf_layout.addWidget(self.supervision_cb)

        layout.addWidget(inf_frame)

        layout.addStretch()

        # Add inputs container to scroll area, then scroll to main layout
        inputs_scroll.setWidget(inputs_container)
        main_layout.addWidget(inputs_scroll, 1)

        self.popup_stack.addWidget(main_container)

    # ================================================================
    # TOGGLE HANDLERS
    # ================================================================
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._generate_reasons_text()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._generate_reasons_text()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._generate_reasons_text()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._generate_reasons_text()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm,
                       self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
                cb.setChecked(False)
        self._generate_reasons_text()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            for cb in [self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual,
                       self.others_hist_stalking, self.others_hist_arson, self.others_curr_violence,
                       self.others_curr_verbal, self.others_curr_sexual, self.others_curr_stalking,
                       self.others_curr_arson]:
                cb.setChecked(False)
        self._generate_reasons_text()

    def _update_reasons_card(self):
        text = self.reasons.text().strip()
        if text:
            preview = text[:100] + "..." if len(text) > 100 else text
            self.cards["reasons"].set_content(preview)
        else:
            self.cards["reasons"].set_content("Click to view/edit reasons")

    def _build_signature_popup(self):
        """Build RC & Signature popup."""
        scroll, layout = self._create_popup_scroll()

        layout.addWidget(self._create_section_label("Responsible Clinician"))

        self.rc_name = self._create_line_edit("RC full name")
        layout.addWidget(self.rc_name)

        self.rc_email = self._create_line_edit("Email (if applicable)")
        layout.addWidget(self.rc_email)

        furnish_lbl = QLabel("I am furnishing this report by:")
        furnish_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151; margin-top: 12px;")
        layout.addWidget(furnish_lbl)

        self.furnish_group = QButtonGroup(self)

        opt1_container = QWidget()
        opt1_layout = QHBoxLayout(opt1_container)
        opt1_layout.setContentsMargins(0, 0, 0, 0)
        opt1_layout.setSpacing(8)

        self.furnish_internal_mail = QRadioButton("Internal mail at")
        self.furnish_internal_mail.setChecked(True)
        self.furnish_internal_mail.setStyleSheet("font-size: 16px;")
        self.furnish_group.addButton(self.furnish_internal_mail)
        opt1_layout.addWidget(self.furnish_internal_mail)

        self.furnish_time = self._create_time_edit()
        self.furnish_time.setFixedWidth(120)
        opt1_layout.addWidget(self.furnish_time)
        opt1_layout.addStretch()
        layout.addWidget(opt1_container)

        self.furnish_electronic = QRadioButton("Electronic communication")
        self.furnish_electronic.setStyleSheet("font-size: 16px;")
        self.furnish_group.addButton(self.furnish_electronic)
        layout.addWidget(self.furnish_electronic)

        self.furnish_other = QRadioButton("Other delivery method")
        self.furnish_other.setStyleSheet("font-size: 16px;")
        self.furnish_group.addButton(self.furnish_other)
        layout.addWidget(self.furnish_other)

        sig_frame = QFrame()
        sig_frame.setStyleSheet("QFrame { background: #f0f9ff; border: none; border-radius: 8px; margin-top: 12px; }")
        sig_layout = QVBoxLayout(sig_frame)
        sig_layout.setContentsMargins(12, 10, 12, 10)
        sig_layout.setSpacing(8)

        sig_header = QLabel("Signature Details")
        sig_header.setStyleSheet("font-size: 16px; font-weight: 700; color: #0369a1;")
        sig_layout.addWidget(sig_header)

        sig_row = QHBoxLayout()
        sig_row.setSpacing(12)

        sig_date_lbl = QLabel("Date:")
        sig_date_lbl.setStyleSheet("font-size: 16px; color: #374151;")
        sig_row.addWidget(sig_date_lbl)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(130)
        sig_row.addWidget(self.rc_sig_date)

        sig_time_lbl = QLabel("Time:")
        sig_time_lbl.setStyleSheet("font-size: 16px; color: #374151;")
        sig_row.addWidget(sig_time_lbl)
        self.rc_sig_time = self._create_time_edit()
        self.rc_sig_time.setFixedWidth(90)
        sig_row.addWidget(self.rc_sig_time)

        sig_row.addStretch()
        sig_layout.addLayout(sig_row)

        layout.addWidget(sig_frame)

        layout.addStretch()

        # Auto-sync to card
        self.rc_sig_date.dateChanged.connect(self._update_signature_card)

        self.popup_stack.addWidget(scroll)

    def _update_signature_card(self):
        parts = []
        if self.rc_name.text().strip():
            parts.append(self.rc_name.text().strip())
        parts.append(self.rc_sig_date.date().toString('dd MMM yyyy'))
        self.cards["signature"].set_content("\n".join(parts) if parts else "Click to enter signature details")

    # ================================================================
    # REASONS TEXT GENERATION
    # ================================================================
    def _generate_reasons_text(self):
        """Generate reasons text based on selected checkboxes."""
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")

        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            eth_simple = ethnicity.replace(" British", "").replace(" Other", "")
            opening_parts.append(eth_simple)

        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        if self.dx_primary.currentText().strip():
            diagnoses.append(self.dx_primary.currentText().strip())
        if self.dx_secondary.currentText().strip():
            diagnoses.append(self.dx_secondary.currentText().strip())

        if diagnoses:
            if opening_parts:
                demo_str = " ".join(opening_parts)
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {joined} which are mental disorders as defined by the Mental Health Act.")
            else:
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} suffers from {joined} which are mental disorders as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree to warrant detention for treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature to warrant detention for treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree to warrant detention for treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    nature_str = " and ".join(nature_types)
                    para1_parts.append(f"The nature of the illness is {nature_str}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        elif opening_parts:
            demo_str = " ".join(opening_parts)
            para1_parts.append(f"{name_display} is a {demo_str} who is currently detained under the Mental Health Act.")

        para2_parts = []
        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("safety of others")

        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")

        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_reasons = []
            if self.poor_compliance_cb.isChecked():
                mh_reasons.append("non compliance")
            if self.limited_insight_cb.isChecked():
                mh_reasons.append("limited insight")
            if mh_reasons:
                reasons_str = "/".join(mh_reasons)
                para2_parts.append(f"Regarding health I would be concerned about {p['pos_l']} mental health deteriorating due to {reasons_str}.")
            else:
                para2_parts.append(f"Regarding health I would be concerned about {p['pos_l']} mental health deteriorating.")

        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health: {details}.")
            else:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health.")

        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            risk_types = [
                ("self neglect", self.self_hist_neglect.isChecked(), self.self_curr_neglect.isChecked()),
                (f"placing of {p['self']} in risky situations", self.self_hist_risky.isChecked(), self.self_curr_risky.isChecked()),
                ("self harm", self.self_hist_harm.isChecked(), self.self_curr_harm.isChecked()),
            ]
            both_items, hist_only, curr_only = [], [], []
            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)
            self_text = f"With respect to {p['pos_l']} own safety I am concerned about"
            parts = []
            if both_items:
                parts.append(f"historical and current {', '.join(both_items)}")
            if hist_only:
                parts.append(f"historical {', '.join(hist_only)}")
            if curr_only:
                parts.append(f"current {', '.join(curr_only)}")
            if parts:
                self_text += " " + ", and ".join(parts) + "."
            else:
                self_text += f" {p['pos_l']} safety."
            para2_parts.append(self_text)

        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            risk_types = [
                ("violence to others", self.others_hist_violence.isChecked(), self.others_curr_violence.isChecked()),
                ("verbal aggression", self.others_hist_verbal.isChecked(), self.others_curr_verbal.isChecked()),
                ("sexual violence", self.others_hist_sexual.isChecked(), self.others_curr_sexual.isChecked()),
                ("stalking", self.others_hist_stalking.isChecked(), self.others_curr_stalking.isChecked()),
                ("arson", self.others_hist_arson.isChecked(), self.others_curr_arson.isChecked()),
            ]
            both_items, hist_only, curr_only = [], [], []
            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)
            others_text = "With respect to risk to others I am concerned about the risk of"
            parts = []
            if both_items:
                parts.append(f"historical and current {', '.join(both_items)}")
            if hist_only:
                parts.append(f"historical {', '.join(hist_only)}")
            if curr_only:
                parts.append(f"current {', '.join(curr_only)}")
            if parts:
                others_text += " " + " and of ".join(parts) + "."
            else:
                others_text += f" {p['pos_l']} potential to cause harm."
            para2_parts.append(others_text)

        para3_parts = []
        if self.tried_failed_cb.isChecked():
            para3_parts.append("Previous attempts at informal admissions have not been successful and I would likewise be concerned about this recurring in this instance hence I do not believe informal admission currently would be appropriate.")
        if self.insight_cb.isChecked():
            para3_parts.append(f"{p['pos']} lack of insight is a significant concern and should {p['subj_l']} be discharged from section, I believe this would significantly impair {p['pos_l']} compliance if informal.")
        if self.compliance_cb.isChecked():
            if para3_parts:
                para3_parts.append(f"Compliance with treatment has also been a significant issue and I do not believe {p['subj_l']} would comply if informal.")
            else:
                para3_parts.append(f"Compliance with treatment has been a significant issue and I do not believe {p['subj_l']} would comply if informal.")
        if self.supervision_cb.isChecked():
            name = patient_name if patient_name else "the patient"
            para3_parts.append(f"I believe {name} needs careful community monitoring under the supervision afforded by the mental health act and I do not believe such supervision would be complied with should {p['subj_l']} remain in the community informally.")

        has_self_risk = self.safety_cb.isChecked() and self.self_harm_cb.isChecked()
        has_others_risk = self.safety_cb.isChecked() and self.others_cb.isChecked()
        if has_self_risk or has_others_risk:
            if has_self_risk and has_others_risk:
                danger = f"other persons and to {p['self']}"
            elif has_others_risk:
                danger = "other persons"
            else:
                danger = p['self']
            para3_parts.append(f"For these reasons, I am of the opinion that if {name_display} were to be discharged, {p['subj_l']} would be likely to act in a manner dangerous to {danger}.")

        paragraphs = []
        if para1_parts:
            paragraphs.append(" ".join(para1_parts))
        if para2_parts:
            paragraphs.append(" ".join(para2_parts))
        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        self.reasons.setText("\n\n".join(paragraphs))
        self._update_reasons_card()

    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.hospital.clear()
            self.nearest_relative.clear()
            self.notice_time.setTime(QTime.currentTime())
            self.notice_date.setDate(QDate.currentDate())
            self.patient_name.clear()
            self.age_spin.setValue(0)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.ethnicity_combo.setCurrentIndex(0)
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.reasons.setText("")
            self.rc_name.clear()
            self.rc_email.clear()
            self.furnish_internal_mail.setChecked(True)
            self.furnish_time.setTime(QTime.currentTime())
            self.rc_sig_date.setDate(QDate.currentDate())
            self.rc_sig_time.setTime(QTime.currentTime())
            self._update_hospital_card()
            self._update_patient_card()
            self._update_reasons_card()
            # Restore my details fields
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form M2",
            f"Form_M2_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return
        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_M2_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form M2 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color (#918C0D) and cream highlight (#FFFED5)
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)
            CREAM_FILL = 'FFFED5'

            # Pre-process: clean ALL paragraphs - remove permission markers, convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), CREAM_FILL)
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), CREAM_FILL)

            def set_entry_box(para, content: str):
                """Set entry box with gold brackets."""
                if not content or not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Opening bracket
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr_ob = bracket_open._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                # Content
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), CREAM_FILL)
                rPr2.append(shd2)
                # Closing bracket
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr_cb = bracket_close._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def format_option_first(para, content: str, strike: bool = False):
                """First option - opening bracket only."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = '['
                    ob = para.runs[0]
                else:
                    ob = para.add_run('[')
                ob.font.name = 'Arial'
                ob.font.size = Pt(12)
                ob.font.bold = True
                ob.font.color.rgb = BRACKET_COLOR
                rPr_ob = ob._element.get_or_add_rPr()
                shd_ob = OxmlElement('w:shd')
                shd_ob.set(qn('w:val'), 'clear')
                shd_ob.set(qn('w:color'), 'auto')
                shd_ob.set(qn('w:fill'), CREAM_FILL)
                rPr_ob.append(shd_ob)
                if strike:
                    ob.font.strike = True
                ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True

            def format_option_middle(para, content: str, strike: bool = False):
                """Middle option - no brackets, just cream highlight."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True

            def format_option_last(para, content: str, strike: bool = False):
                """Last option - closing bracket only."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                if strike:
                    ct.font.strike = True
                cb = para.add_run(']')
                cb.font.name = 'Arial'
                cb.font.size = Pt(12)
                cb.font.bold = True
                cb.font.color.rgb = BRACKET_COLOR
                rPr_cb = cb._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)
                if strike:
                    cb.font.strike = True

            def format_cream_para(para, content: str):
                """Cream highlight paragraph."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = content
                    ct = para.runs[0]
                else:
                    ct = para.add_run(content)
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)

            def format_closing_bracket_para(para):
                """Paragraph with cream and closing bracket."""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                if para.runs:
                    para.runs[0].text = '                                                                                                    '
                    ct = para.runs[0]
                else:
                    ct = para.add_run('                                                                                                    ')
                ct.font.name = 'Arial'
                ct.font.size = Pt(12)
                rPr_ct = ct._element.get_or_add_rPr()
                shd_ct = OxmlElement('w:shd')
                shd_ct.set(qn('w:val'), 'clear')
                shd_ct.set(qn('w:color'), 'auto')
                shd_ct.set(qn('w:fill'), CREAM_FILL)
                rPr_ct.append(shd_ct)
                cb = para.add_run(']')
                cb.font.name = 'Arial'
                cb.font.size = Pt(12)
                cb.font.bold = True
                cb.font.color.rgb = BRACKET_COLOR
                rPr_cb = cb._element.get_or_add_rPr()
                shd_cb = OxmlElement('w:shd')
                shd_cb.set(qn('w:val'), 'clear')
                shd_cb.set(qn('w:color'), 'auto')
                shd_cb.set(qn('w:fill'), CREAM_FILL)
                rPr_cb.append(shd_cb)

            def format_sig_line(para, label1, label2, value1="", value2=""):
                """Format signature line: Label1 [value1] Label2 [value2]"""
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                pPr = para._element.get_or_add_pPr()
                for old_shd in pPr.findall(qn('w:shd')):
                    pPr.remove(old_shd)
                # Label1
                if para.runs:
                    para.runs[0].text = label1
                    l1 = para.runs[0]
                else:
                    l1 = para.add_run(label1)
                l1.font.name = 'Arial'
                l1.font.size = Pt(12)
                # First placeholder
                ob1 = para.add_run('[')
                ob1.font.bold = True
                ob1.font.color.rgb = BRACKET_COLOR
                rPr_ob1 = ob1._element.get_or_add_rPr()
                shd_ob1 = OxmlElement('w:shd')
                shd_ob1.set(qn('w:val'), 'clear')
                shd_ob1.set(qn('w:color'), 'auto')
                shd_ob1.set(qn('w:fill'), CREAM_FILL)
                rPr_ob1.append(shd_ob1)
                content1 = value1 if value1 else '                              '
                c1 = para.add_run(content1)
                rPr_c1 = c1._element.get_or_add_rPr()
                shd_c1 = OxmlElement('w:shd')
                shd_c1.set(qn('w:val'), 'clear')
                shd_c1.set(qn('w:color'), 'auto')
                shd_c1.set(qn('w:fill'), CREAM_FILL)
                rPr_c1.append(shd_c1)
                cb1 = para.add_run(']')
                cb1.font.bold = True
                cb1.font.color.rgb = BRACKET_COLOR
                rPr_cb1 = cb1._element.get_or_add_rPr()
                shd_cb1 = OxmlElement('w:shd')
                shd_cb1.set(qn('w:val'), 'clear')
                shd_cb1.set(qn('w:color'), 'auto')
                shd_cb1.set(qn('w:fill'), CREAM_FILL)
                rPr_cb1.append(shd_cb1)
                # Label2
                l2 = para.add_run(' ' + label2)
                l2.font.name = 'Arial'
                l2.font.size = Pt(12)
                # Second placeholder
                ob2 = para.add_run('[')
                ob2.font.bold = True
                ob2.font.color.rgb = BRACKET_COLOR
                rPr_ob2 = ob2._element.get_or_add_rPr()
                shd_ob2 = OxmlElement('w:shd')
                shd_ob2.set(qn('w:val'), 'clear')
                shd_ob2.set(qn('w:color'), 'auto')
                shd_ob2.set(qn('w:fill'), CREAM_FILL)
                rPr_ob2.append(shd_ob2)
                content2 = value2 if value2 else '                              '
                c2 = para.add_run(content2)
                rPr_c2 = c2._element.get_or_add_rPr()
                shd_c2 = OxmlElement('w:shd')
                shd_c2.set(qn('w:val'), 'clear')
                shd_c2.set(qn('w:color'), 'auto')
                shd_c2.set(qn('w:fill'), CREAM_FILL)
                rPr_c2.append(shd_c2)
                cb2 = para.add_run(']')
                cb2.font.bold = True
                cb2.font.color.rgb = BRACKET_COLOR
                rPr_cb2 = cb2._element.get_or_add_rPr()
                shd_cb2 = OxmlElement('w:shd')
                shd_cb2.set(qn('w:val'), 'clear')
                shd_cb2.set(qn('w:color'), 'auto')
                shd_cb2.set(qn('w:fill'), CREAM_FILL)
                rPr_cb2.append(shd_cb2)

            paragraphs = doc.paragraphs

            # PART 1 - Page 1 entry boxes
            # Hospital (para 5)
            hospital_text = self.hospital.text().strip()
            set_entry_box(paragraphs[5], hospital_text)

            # Nearest relative (para 7)
            nr_text = self.nearest_relative.text().strip()
            set_entry_box(paragraphs[7], nr_text)

            # Notice time (para 9)
            notice_time = self.notice_time.time().toString("HH:mm")
            set_entry_box(paragraphs[9], notice_time)

            # Notice date (para 11)
            notice_date = self.notice_date.date().toString("dd MMMM yyyy")
            set_entry_box(paragraphs[11], notice_date)

            # Patient name (para 13)
            patient_text = self.patient_name.text().strip()
            set_entry_box(paragraphs[13], patient_text)

            # Reasons box (para 16-18)
            reasons_text = self.reasons.text().strip()
            if not reasons_text:
                reasons_text = '                                                                                                    '

            # Para 16 - opening bracket + content
            for run in paragraphs[16].runs:
                run.text = ""
            while len(paragraphs[16].runs) > 1:
                paragraphs[16]._element.remove(paragraphs[16].runs[-1]._element)
            pPr16 = paragraphs[16]._element.get_or_add_pPr()
            for old_shd in pPr16.findall(qn('w:shd')):
                pPr16.remove(old_shd)
            if paragraphs[16].runs:
                paragraphs[16].runs[0].text = '['
                ob16 = paragraphs[16].runs[0]
            else:
                ob16 = paragraphs[16].add_run('[')
            ob16.font.name = 'Arial'
            ob16.font.size = Pt(12)
            ob16.font.bold = True
            ob16.font.color.rgb = BRACKET_COLOR
            rPr_ob16 = ob16._element.get_or_add_rPr()
            shd_ob16 = OxmlElement('w:shd')
            shd_ob16.set(qn('w:val'), 'clear')
            shd_ob16.set(qn('w:color'), 'auto')
            shd_ob16.set(qn('w:fill'), CREAM_FILL)
            rPr_ob16.append(shd_ob16)
            ct16 = paragraphs[16].add_run(reasons_text)
            ct16.font.name = 'Arial'
            ct16.font.size = Pt(12)
            rPr_ct16 = ct16._element.get_or_add_rPr()
            shd_ct16 = OxmlElement('w:shd')
            shd_ct16.set(qn('w:val'), 'clear')
            shd_ct16.set(qn('w:color'), 'auto')
            shd_ct16.set(qn('w:fill'), CREAM_FILL)
            rPr_ct16.append(shd_ct16)

            # Para 17 - cream continuation
            format_cream_para(paragraphs[17], '                                                                                                    ')

            # Para 18 - cream + closing bracket
            format_closing_bracket_para(paragraphs[18])

            # Para 19: "[If you need to continue on a separate sheet please indicate here [ ] and attach that sheet to this form]"
            for run in paragraphs[19].runs:
                run.text = ""
            while len(paragraphs[19].runs) > 1:
                paragraphs[19]._element.remove(paragraphs[19].runs[-1]._element)
            pPr19 = paragraphs[19]._element.get_or_add_pPr()
            for old_shd in pPr19.findall(qn('w:shd')):
                pPr19.remove(old_shd)
            if paragraphs[19].runs:
                paragraphs[19].runs[0].text = '[If you need to continue on a separate sheet please indicate here'
                t19a = paragraphs[19].runs[0]
            else:
                t19a = paragraphs[19].add_run('[If you need to continue on a separate sheet please indicate here')
            t19a.font.name = 'Arial'
            t19a.font.size = Pt(12)
            # Gold bracket open
            ob19 = paragraphs[19].add_run('[')
            ob19.font.bold = True
            ob19.font.color.rgb = BRACKET_COLOR
            rPr_ob19 = ob19._element.get_or_add_rPr()
            shd_ob19 = OxmlElement('w:shd')
            shd_ob19.set(qn('w:val'), 'clear')
            shd_ob19.set(qn('w:color'), 'auto')
            shd_ob19.set(qn('w:fill'), CREAM_FILL)
            rPr_ob19.append(shd_ob19)
            # Placeholder content
            c19 = paragraphs[19].add_run('     ')
            rPr_c19 = c19._element.get_or_add_rPr()
            shd_c19 = OxmlElement('w:shd')
            shd_c19.set(qn('w:val'), 'clear')
            shd_c19.set(qn('w:color'), 'auto')
            shd_c19.set(qn('w:fill'), CREAM_FILL)
            rPr_c19.append(shd_c19)
            # Gold bracket close
            cb19 = paragraphs[19].add_run(']')
            cb19.font.bold = True
            cb19.font.color.rgb = BRACKET_COLOR
            rPr_cb19 = cb19._element.get_or_add_rPr()
            shd_cb19 = OxmlElement('w:shd')
            shd_cb19.set(qn('w:val'), 'clear')
            shd_cb19.set(qn('w:color'), 'auto')
            shd_cb19.set(qn('w:fill'), CREAM_FILL)
            rPr_cb19.append(shd_cb19)
            # Rest of text
            t19b = paragraphs[19].add_run('and attach that sheet to this form]')
            t19b.font.name = 'Arial'
            t19b.font.size = Pt(12)

            # PART 1 Furnishing options (para 21-24)
            any_furnish_selected = self.furnish_internal_mail.isChecked() or self.furnish_electronic.isChecked() or self.furnish_other.isChecked()
            strike_internal = any_furnish_selected and not self.furnish_internal_mail.isChecked()
            strike_electronic = any_furnish_selected and not self.furnish_electronic.isChecked()
            strike_other = any_furnish_selected and not self.furnish_other.isChecked()

            # Para 21 - internal mail with [time] text
            format_option_first(paragraphs[21], "consigning it to the hospital managers' internal mail system today at [time].", strike=strike_internal)

            # Para 22 - time placeholder with golden brackets
            furnish_time = self.furnish_time.time().toString("HH:mm")
            furnish_time_content = furnish_time if furnish_time else '                              '
            for run in paragraphs[22].runs:
                run.text = ""
            while len(paragraphs[22].runs) > 1:
                paragraphs[22]._element.remove(paragraphs[22].runs[-1]._element)
            pPr22 = paragraphs[22]._element.get_or_add_pPr()
            for old_shd in pPr22.findall(qn('w:shd')):
                pPr22.remove(old_shd)
            if paragraphs[22].runs:
                paragraphs[22].runs[0].text = '['
                ob22 = paragraphs[22].runs[0]
            else:
                ob22 = paragraphs[22].add_run('[')
            ob22.font.name = 'Arial'
            ob22.font.size = Pt(12)
            ob22.font.bold = True
            ob22.font.color.rgb = BRACKET_COLOR
            if strike_internal:
                ob22.font.strike = True
            rPr_ob22 = ob22._element.get_or_add_rPr()
            shd_ob22 = OxmlElement('w:shd')
            shd_ob22.set(qn('w:val'), 'clear')
            shd_ob22.set(qn('w:color'), 'auto')
            shd_ob22.set(qn('w:fill'), CREAM_FILL)
            rPr_ob22.append(shd_ob22)
            c22 = paragraphs[22].add_run(furnish_time_content)
            c22.font.name = 'Arial'
            c22.font.size = Pt(12)
            if strike_internal:
                c22.font.strike = True
            rPr_c22 = c22._element.get_or_add_rPr()
            shd_c22 = OxmlElement('w:shd')
            shd_c22.set(qn('w:val'), 'clear')
            shd_c22.set(qn('w:color'), 'auto')
            shd_c22.set(qn('w:fill'), CREAM_FILL)
            rPr_c22.append(shd_c22)
            cb22 = paragraphs[22].add_run(']')
            cb22.font.name = 'Arial'
            cb22.font.size = Pt(12)
            cb22.font.bold = True
            cb22.font.color.rgb = BRACKET_COLOR
            if strike_internal:
                cb22.font.strike = True
            rPr_cb22 = cb22._element.get_or_add_rPr()
            shd_cb22 = OxmlElement('w:shd')
            shd_cb22.set(qn('w:val'), 'clear')
            shd_cb22.set(qn('w:color'), 'auto')
            shd_cb22.set(qn('w:fill'), CREAM_FILL)
            rPr_cb22.append(shd_cb22)

            # Para 23 - electronic communication
            format_option_middle(paragraphs[23], "today sending it to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.", strike=strike_electronic)

            # Para 24 - other delivery (closing bracket)
            format_option_last(paragraphs[24], "sending or delivering it without using the hospital managers' internal mail system.", strike=strike_other)

            # PART 1 Signature section
            # Para 25: Signed [ ] Responsible clinician
            rc_name = self.rc_name.text().strip()
            for run in paragraphs[25].runs:
                run.text = ""
            while len(paragraphs[25].runs) > 1:
                paragraphs[25]._element.remove(paragraphs[25].runs[-1]._element)
            pPr25 = paragraphs[25]._element.get_or_add_pPr()
            for old_shd in pPr25.findall(qn('w:shd')):
                pPr25.remove(old_shd)
            if paragraphs[25].runs:
                paragraphs[25].runs[0].text = 'Signed '
                sl25 = paragraphs[25].runs[0]
            else:
                sl25 = paragraphs[25].add_run('Signed ')
            sl25.font.name = 'Arial'
            sl25.font.size = Pt(12)
            ob25 = paragraphs[25].add_run('[')
            ob25.font.bold = True
            ob25.font.color.rgb = BRACKET_COLOR
            rPr_ob25 = ob25._element.get_or_add_rPr()
            shd_ob25 = OxmlElement('w:shd')
            shd_ob25.set(qn('w:val'), 'clear')
            shd_ob25.set(qn('w:color'), 'auto')
            shd_ob25.set(qn('w:fill'), CREAM_FILL)
            rPr_ob25.append(shd_ob25)
            c25 = paragraphs[25].add_run('                                        ')
            rPr_c25 = c25._element.get_or_add_rPr()
            shd_c25 = OxmlElement('w:shd')
            shd_c25.set(qn('w:val'), 'clear')
            shd_c25.set(qn('w:color'), 'auto')
            shd_c25.set(qn('w:fill'), CREAM_FILL)
            rPr_c25.append(shd_c25)
            cb25 = paragraphs[25].add_run(']')
            cb25.font.bold = True
            cb25.font.color.rgb = BRACKET_COLOR
            rPr_cb25 = cb25._element.get_or_add_rPr()
            shd_cb25 = OxmlElement('w:shd')
            shd_cb25.set(qn('w:val'), 'clear')
            shd_cb25.set(qn('w:color'), 'auto')
            shd_cb25.set(qn('w:fill'), CREAM_FILL)
            rPr_cb25.append(shd_cb25)
            suf25 = paragraphs[25].add_run(' Responsible clinician')
            suf25.font.name = 'Arial'
            suf25.font.size = Pt(12)

            # Para 26: PRINT NAME [ rc_name ]
            for run in paragraphs[26].runs:
                run.text = ""
            while len(paragraphs[26].runs) > 1:
                paragraphs[26]._element.remove(paragraphs[26].runs[-1]._element)
            pPr26 = paragraphs[26]._element.get_or_add_pPr()
            for old_shd in pPr26.findall(qn('w:shd')):
                pPr26.remove(old_shd)
            if paragraphs[26].runs:
                paragraphs[26].runs[0].text = 'PRINT NAME '
                l26 = paragraphs[26].runs[0]
            else:
                l26 = paragraphs[26].add_run('PRINT NAME ')
            l26.font.name = 'Arial'
            l26.font.size = Pt(12)
            ob26 = paragraphs[26].add_run('[')
            ob26.font.bold = True
            ob26.font.color.rgb = BRACKET_COLOR
            rPr_ob26 = ob26._element.get_or_add_rPr()
            shd_ob26 = OxmlElement('w:shd')
            shd_ob26.set(qn('w:val'), 'clear')
            shd_ob26.set(qn('w:color'), 'auto')
            shd_ob26.set(qn('w:fill'), CREAM_FILL)
            rPr_ob26.append(shd_ob26)
            name_content = rc_name if rc_name else '                                        '
            c26 = paragraphs[26].add_run(name_content)
            rPr_c26 = c26._element.get_or_add_rPr()
            shd_c26 = OxmlElement('w:shd')
            shd_c26.set(qn('w:val'), 'clear')
            shd_c26.set(qn('w:color'), 'auto')
            shd_c26.set(qn('w:fill'), CREAM_FILL)
            rPr_c26.append(shd_c26)
            cb26 = paragraphs[26].add_run(']')
            cb26.font.bold = True
            cb26.font.color.rgb = BRACKET_COLOR
            rPr_cb26 = cb26._element.get_or_add_rPr()
            shd_cb26 = OxmlElement('w:shd')
            shd_cb26.set(qn('w:val'), 'clear')
            shd_cb26.set(qn('w:color'), 'auto')
            shd_cb26.set(qn('w:fill'), CREAM_FILL)
            rPr_cb26.append(shd_cb26)

            # Para 27: Email address (if applicable) [ rc_email ]
            rc_email = self.rc_email.text().strip()
            for run in paragraphs[27].runs:
                run.text = ""
            while len(paragraphs[27].runs) > 1:
                paragraphs[27]._element.remove(paragraphs[27].runs[-1]._element)
            pPr27 = paragraphs[27]._element.get_or_add_pPr()
            for old_shd in pPr27.findall(qn('w:shd')):
                pPr27.remove(old_shd)
            if paragraphs[27].runs:
                paragraphs[27].runs[0].text = 'Email address (if applicable) '
                l27 = paragraphs[27].runs[0]
            else:
                l27 = paragraphs[27].add_run('Email address (if applicable) ')
            l27.font.name = 'Arial'
            l27.font.size = Pt(12)
            ob27 = paragraphs[27].add_run('[')
            ob27.font.bold = True
            ob27.font.color.rgb = BRACKET_COLOR
            rPr_ob27 = ob27._element.get_or_add_rPr()
            shd_ob27 = OxmlElement('w:shd')
            shd_ob27.set(qn('w:val'), 'clear')
            shd_ob27.set(qn('w:color'), 'auto')
            shd_ob27.set(qn('w:fill'), CREAM_FILL)
            rPr_ob27.append(shd_ob27)
            email_content = rc_email if rc_email else '                                        '
            c27 = paragraphs[27].add_run(email_content)
            rPr_c27 = c27._element.get_or_add_rPr()
            shd_c27 = OxmlElement('w:shd')
            shd_c27.set(qn('w:val'), 'clear')
            shd_c27.set(qn('w:color'), 'auto')
            shd_c27.set(qn('w:fill'), CREAM_FILL)
            rPr_c27.append(shd_c27)
            cb27 = paragraphs[27].add_run(']')
            cb27.font.bold = True
            cb27.font.color.rgb = BRACKET_COLOR
            rPr_cb27 = cb27._element.get_or_add_rPr()
            shd_cb27 = OxmlElement('w:shd')
            shd_cb27.set(qn('w:val'), 'clear')
            shd_cb27.set(qn('w:color'), 'auto')
            shd_cb27.set(qn('w:fill'), CREAM_FILL)
            rPr_cb27.append(shd_cb27)

            # Para 28: Date [ date ] Time [ time ]
            sig_date = self.rc_sig_date.date().toString("dd MMMM yyyy")
            sig_time = self.rc_sig_time.time().toString("HH:mm")
            format_sig_line(paragraphs[28], "Date ", "Time ", sig_date, sig_time)

            # PART 2 Receipt options (para 32-37) - NO strikethrough in Part 2
            # Para 32: internal mail - opening bracket
            format_option_first(paragraphs[32], "furnished to the hospital managers through their internal mail system.", strike=False)

            # Para 33: electronic communication - middle
            format_option_middle(paragraphs[33], "furnished to the hospital managers, or a person authorised by them to receive it, by means of electronic communication.", strike=False)

            # Para 34: received at [time] - middle
            format_option_middle(paragraphs[34], "received by me on behalf of the hospital managers at [time]", strike=False)

            # Para 35: time placeholder with golden brackets
            for run in paragraphs[35].runs:
                run.text = ""
            while len(paragraphs[35].runs) > 1:
                paragraphs[35]._element.remove(paragraphs[35].runs[-1]._element)
            pPr35 = paragraphs[35]._element.get_or_add_pPr()
            for old_shd in pPr35.findall(qn('w:shd')):
                pPr35.remove(old_shd)
            if paragraphs[35].runs:
                paragraphs[35].runs[0].text = '['
                ob35 = paragraphs[35].runs[0]
            else:
                ob35 = paragraphs[35].add_run('[')
            ob35.font.name = 'Arial'
            ob35.font.size = Pt(12)
            ob35.font.bold = True
            ob35.font.color.rgb = BRACKET_COLOR
            rPr_ob35 = ob35._element.get_or_add_rPr()
            shd_ob35 = OxmlElement('w:shd')
            shd_ob35.set(qn('w:val'), 'clear')
            shd_ob35.set(qn('w:color'), 'auto')
            shd_ob35.set(qn('w:fill'), CREAM_FILL)
            rPr_ob35.append(shd_ob35)
            c35 = paragraphs[35].add_run('                                                      ')
            c35.font.name = 'Arial'
            c35.font.size = Pt(12)
            rPr_c35 = c35._element.get_or_add_rPr()
            shd_c35 = OxmlElement('w:shd')
            shd_c35.set(qn('w:val'), 'clear')
            shd_c35.set(qn('w:color'), 'auto')
            shd_c35.set(qn('w:fill'), CREAM_FILL)
            rPr_c35.append(shd_c35)
            cb35 = paragraphs[35].add_run(']')
            cb35.font.name = 'Arial'
            cb35.font.size = Pt(12)
            cb35.font.bold = True
            cb35.font.color.rgb = BRACKET_COLOR
            rPr_cb35 = cb35._element.get_or_add_rPr()
            shd_cb35 = OxmlElement('w:shd')
            shd_cb35.set(qn('w:val'), 'clear')
            shd_cb35.set(qn('w:color'), 'auto')
            shd_cb35.set(qn('w:fill'), CREAM_FILL)
            rPr_cb35.append(shd_cb35)

            # Para 36: on [date]. - middle
            format_option_middle(paragraphs[36], "on [date].", strike=False)

            # Para 37: date placeholder with closing bracket
            format_closing_bracket_para(paragraphs[37])

            # Para 38: Signed [ ] on behalf of hospital managers
            for run in paragraphs[38].runs:
                run.text = ""
            while len(paragraphs[38].runs) > 1:
                paragraphs[38]._element.remove(paragraphs[38].runs[-1]._element)
            pPr38 = paragraphs[38]._element.get_or_add_pPr()
            for old_shd in pPr38.findall(qn('w:shd')):
                pPr38.remove(old_shd)
            if paragraphs[38].runs:
                paragraphs[38].runs[0].text = 'Signed '
                sl38 = paragraphs[38].runs[0]
            else:
                sl38 = paragraphs[38].add_run('Signed ')
            sl38.font.name = 'Arial'
            sl38.font.size = Pt(12)
            ob38 = paragraphs[38].add_run('[')
            ob38.font.bold = True
            ob38.font.color.rgb = BRACKET_COLOR
            rPr_ob38 = ob38._element.get_or_add_rPr()
            shd_ob38 = OxmlElement('w:shd')
            shd_ob38.set(qn('w:val'), 'clear')
            shd_ob38.set(qn('w:color'), 'auto')
            shd_ob38.set(qn('w:fill'), CREAM_FILL)
            rPr_ob38.append(shd_ob38)
            c38 = paragraphs[38].add_run('                                        ')
            rPr_c38 = c38._element.get_or_add_rPr()
            shd_c38 = OxmlElement('w:shd')
            shd_c38.set(qn('w:val'), 'clear')
            shd_c38.set(qn('w:color'), 'auto')
            shd_c38.set(qn('w:fill'), CREAM_FILL)
            rPr_c38.append(shd_c38)
            cb38 = paragraphs[38].add_run(']')
            cb38.font.bold = True
            cb38.font.color.rgb = BRACKET_COLOR
            rPr_cb38 = cb38._element.get_or_add_rPr()
            shd_cb38 = OxmlElement('w:shd')
            shd_cb38.set(qn('w:val'), 'clear')
            shd_cb38.set(qn('w:color'), 'auto')
            shd_cb38.set(qn('w:fill'), CREAM_FILL)
            rPr_cb38.append(shd_cb38)
            suf38 = paragraphs[38].add_run(' on behalf of the hospital managers')
            suf38.font.name = 'Arial'
            suf38.font.size = Pt(12)

            # Para 39: PRINT NAME [ ] Date [ ]
            format_sig_line(paragraphs[39], "PRINT NAME ", "Date ", "", "")

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form M2 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        furnish_method = "internal_mail"
        if self.furnish_electronic.isChecked():
            furnish_method = "electronic"
        elif self.furnish_other.isChecked():
            furnish_method = "other"
        gender = None
        if self.gender_male.isChecked():
            gender = "male"
        elif self.gender_female.isChecked():
            gender = "female"
        elif self.gender_other.isChecked():
            gender = "other"
        return {
            "hospital": self.hospital.text(),
            "nearest_relative": self.nearest_relative.text(),
            "notice_time": self.notice_time.time().toString("HH:mm"),
            "notice_date": self.notice_date.date().toString("yyyy-MM-dd"),
            "patient_name": self.patient_name.text(),
            "age": self.age_spin.value(),
            "gender": gender,
            "ethnicity": self.ethnicity_combo.currentText(),
            "dx_primary": self.dx_primary.currentText(),
            "dx_secondary": self.dx_secondary.currentText(),
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_slider": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_hist_neglect": self.self_hist_neglect.isChecked(),
            "self_hist_risky": self.self_hist_risky.isChecked(),
            "self_hist_harm": self.self_hist_harm.isChecked(),
            "self_curr_neglect": self.self_curr_neglect.isChecked(),
            "self_curr_risky": self.self_curr_risky.isChecked(),
            "self_curr_harm": self.self_curr_harm.isChecked(),
            "others": self.others_cb.isChecked(),
            "others_hist_violence": self.others_hist_violence.isChecked(),
            "others_hist_verbal": self.others_hist_verbal.isChecked(),
            "others_hist_sexual": self.others_hist_sexual.isChecked(),
            "others_hist_stalking": self.others_hist_stalking.isChecked(),
            "others_hist_arson": self.others_hist_arson.isChecked(),
            "others_curr_violence": self.others_curr_violence.isChecked(),
            "others_curr_verbal": self.others_curr_verbal.isChecked(),
            "others_curr_sexual": self.others_curr_sexual.isChecked(),
            "others_curr_stalking": self.others_curr_stalking.isChecked(),
            "others_curr_arson": self.others_curr_arson.isChecked(),
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "reasons": self.reasons.text(),
            "rc_name": self.rc_name.text(),
            "rc_email": self.rc_email.text(),
            "furnish_method": furnish_method,
            "furnish_time": self.furnish_time.time().toString("HH:mm"),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "rc_sig_time": self.rc_sig_time.time().toString("HH:mm"),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.hospital.setText(state.get("hospital", ""))
        self.nearest_relative.setText(state.get("nearest_relative", ""))
        if state.get("notice_time"):
            self.notice_time.setTime(QTime.fromString(state["notice_time"], "HH:mm"))
        if state.get("notice_date"):
            self.notice_date.setDate(QDate.fromString(state["notice_date"], "yyyy-MM-dd"))
        self.patient_name.setText(state.get("patient_name", ""))
        self.age_spin.setValue(state.get("age", 0))
        gender = state.get("gender")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        ethnicity = state.get("ethnicity", "Not specified")
        idx = self.ethnicity_combo.findText(ethnicity)
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)
        dx_secondary = state.get("dx_secondary", "")
        idx = self.dx_secondary.findText(dx_secondary)
        if idx >= 0:
            self.dx_secondary.setCurrentIndex(idx)
        else:
            self.dx_secondary.setCurrentText(dx_secondary)
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_slider", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_hist_neglect.setChecked(state.get("self_hist_neglect", False))
        self.self_hist_risky.setChecked(state.get("self_hist_risky", False))
        self.self_hist_harm.setChecked(state.get("self_hist_harm", False))
        self.self_curr_neglect.setChecked(state.get("self_curr_neglect", False))
        self.self_curr_risky.setChecked(state.get("self_curr_risky", False))
        self.self_curr_harm.setChecked(state.get("self_curr_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.others_hist_violence.setChecked(state.get("others_hist_violence", False))
        self.others_hist_verbal.setChecked(state.get("others_hist_verbal", False))
        self.others_hist_sexual.setChecked(state.get("others_hist_sexual", False))
        self.others_hist_stalking.setChecked(state.get("others_hist_stalking", False))
        self.others_hist_arson.setChecked(state.get("others_hist_arson", False))
        self.others_curr_violence.setChecked(state.get("others_curr_violence", False))
        self.others_curr_verbal.setChecked(state.get("others_curr_verbal", False))
        self.others_curr_sexual.setChecked(state.get("others_curr_sexual", False))
        self.others_curr_stalking.setChecked(state.get("others_curr_stalking", False))
        self.others_curr_arson.setChecked(state.get("others_curr_arson", False))
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        self.reasons.setText(state.get("reasons", ""))
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_email.setText(state.get("rc_email", ""))
        furnish_method = state.get("furnish_method", "internal_mail")
        if furnish_method == "electronic":
            self.furnish_electronic.setChecked(True)
        elif furnish_method == "other":
            self.furnish_other.setChecked(True)
        else:
            self.furnish_internal_mail.setChecked(True)
        if state.get("furnish_time"):
            self.furnish_time.setTime(QTime.fromString(state["furnish_time"], "HH:mm"))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        if state.get("rc_sig_time"):
            self.rc_sig_time.setTime(QTime.fromString(state["rc_sig_time"], "HH:mm"))
        self._update_hospital_card()
        self._update_patient_card()
        self._update_reasons_card()

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[M2Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[M2Form] Set gender: {gender}")
        # Set age if not already set (from age field or calculate from DOB)
        if hasattr(self, 'age_spin') and self.age_spin.value() == 0:
            age = patient_info.get("age")
            if not age and patient_info.get("dob"):
                from datetime import datetime
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if age and 0 < age < 120:
                self.age_spin.setValue(age)
                print(f"[M2Form] Set age: {age}")
        # Set ethnicity if not already selected
        if patient_info.get("ethnicity") and hasattr(self, 'ethnicity_combo'):
            current = self.ethnicity_combo.currentText()
            if current in ("Ethnicity", "Not specified", "Select..."):
                idx = self.ethnicity_combo.findText(patient_info["ethnicity"], Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    self.ethnicity_combo.setCurrentIndex(idx)
                    print(f"[M2Form] Set ethnicity: {patient_info['ethnicity']}")
