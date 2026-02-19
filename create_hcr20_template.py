#!/usr/bin/env python3
"""
Create a blank HCR-20 V3 Assessment Report template matching the provided format.
"""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE

def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="{top}" w:sz="{top_sz}" w:color="{top_color}"/>'
        r'<w:left w:val="{left}" w:sz="{left_sz}" w:color="{left_color}"/>'
        r'<w:bottom w:val="{bottom}" w:sz="{bottom_sz}" w:color="{bottom_color}"/>'
        r'<w:right w:val="{right}" w:sz="{right_sz}" w:color="{right_color}"/>'
        r'</w:tcBorders>'.format(**kwargs)
    )
    tcPr.append(tcBorders)

def add_double_border_paragraph(doc, text, bold=False, font_size=11, font_name='Arial Narrow'):
    """Add paragraph with double border styling (simulated through table)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    return para

def create_header_info_table(doc):
    """Create the header information box with double border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)

    # Set double border on cell
    set_cell_border(
        cell,
        top="double", top_sz="4", top_color="000000",
        bottom="double", bottom_sz="4", bottom_color="000000",
        left="double", left_sz="4", left_color="000000",
        right="double", right_sz="4", right_color="000000"
    )

    # Add content to the cell
    fields = [
        ("NAME:", "[Patient Name]"),
        ("", ""),
        ("D.O.B:", "[Date of Birth]"),
        ("", ""),
        ("AGE:", "[Age]"),
        ("", ""),
        ("NHS NUMBER:", "[NHS Number]"),
        ("", ""),
        ("ADDRESS:", "[Address Line 1]"),
        ("", "[Address Line 2]"),
        ("", ""),
        ("DATE OF ADMISSION:", "[Date]"),
        ("", ""),
        ("LEGAL STATUS:", "[Legal Status]"),
        ("", ""),
        ("STRUCTURED CLINICAL JUDGMENT", ""),
        ("TOOLS INCLUDED IN THIS REPORT:", "HCR-20 V3"),
        ("", ""),
        ("AUTHOR OF ORIGINAL REPORT:", "[Name (Title)]"),
        ("", ""),
        ("AUTHOR OF UPDATE REPORTS:", "[Name (Title)]"),
        ("Under the Supervision of:", "[Name (Title)]"),
        ("", ""),
        ("REPORT SENT FOR REVIEW TO:", "[MDT Name]"),
        ("", ""),
        ("Date of original report:", "[Month Year]"),
        ("", ""),
        ("Date of update report:", "[Month Year]"),
        ("", ""),
        ("Date next update due:", "[Month Year]"),
    ]

    for i, (label, value) in enumerate(fields):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()

        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

        if label:
            run = para.add_run(label + "\t")
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            run.bold = True

        if value:
            run = para.add_run(value)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            run.bold = False

    return table

def add_item_header(doc, item_code, item_title):
    """Add an HCR-20 item header."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)

    run = para.add_run(f"ITEM {item_code}")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run(item_title)
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True

    return para

def add_content_placeholder(doc, placeholder_text="[Enter content here]"):
    """Add a placeholder paragraph for content."""
    para = doc.add_paragraph()
    run = para.add_run(placeholder_text)
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.italic = True
    return para

def add_presence_relevance_table(doc):
    """Add the Presence/Relevance rating table for each item."""
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Row 1: Presence
    table.cell(0, 0).text = "Presence"
    table.cell(0, 1).text = "[Absent / Partially Present / Present]"

    # Row 2: Relevance
    table.cell(1, 0).text = "Relevance"
    table.cell(1, 1).text = "[Low / Moderate / High relevance]"

    # Format cells
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)

    # Bold the labels
    for row in table.rows:
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()
    return table

def add_subsection(doc, title):
    """Add a subsection header."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    run.italic = True
    return para

def create_hcr20_template():
    """Create the complete HCR-20 V3 template."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(11)

    # ===== TITLE =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HCR-20 Assessment Report")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    # ===== HEADER INFO BOX =====
    create_header_info_table(doc)

    doc.add_paragraph()

    # ===== CONFIDENTIALITY NOTICE =====
    para = doc.add_paragraph()
    run = para.add_run("This report is CONFIDENTIAL and should be restricted to persons with involvement with the patient. Sections of this report should not be cut and pasted into future reports without the expressed permission of the author.")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(10)
    run.italic = True

    # Page break
    doc.add_page_break()

    # ===== HCR-20 V3 INTRODUCTION =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HCR-20 V3: RISK ASSESSMENT REPORT")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()

    # Introduction text
    para = doc.add_paragraph()
    run = para.add_run("HCR-20 V3 (Douglas et al., 2013)")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run("The HCR-20 V3 is comprised of 20 risk factors across three subscales: ")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run = para.add_run("Historical")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(", ")
    run = para.add_run("Clinical")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(", and ")
    run = para.add_run("Risk Management")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    para.add_run(".")

    para = doc.add_paragraph()
    run = para.add_run("The presence of factors is coded using a 3-level response format:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)

    # Presence ratings list
    for rating in ["N = absent", "P = possibly or partially present", "Y = definitely present"]:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        run = para.add_run("• " + rating)
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)

    para = doc.add_paragraph()
    run = para.add_run('"Omit" is used when there is no reliable information by which to judge the presence of the factor.')
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)

    para = doc.add_paragraph()
    run = para.add_run('Evaluators then assess whether each risk factor is "relevant" to an individual\'s risk for violent behaviour. This decision is based on whether the factor:')
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.5)
    run = para.add_run("• Is likely to influence the individual's decision to act in a violent manner in the future")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)

    para = doc.add_paragraph()
    run = para.add_run("The relevance of each factor is defined as '")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run = para.add_run("Low")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    para.add_run("', '")
    run = para.add_run("Moderate")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    para.add_run("' or '")
    run = para.add_run("High")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    para.add_run("' for Historical and Clinical items, and Risk Management items.")

    doc.add_paragraph()

    # ===== SOURCES OF INFORMATION =====
    para = doc.add_paragraph()
    run = para.add_run("Sources of information:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    # Sources table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.text = "[Source]"
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # ===== HISTORICAL ITEMS =====
    section_title = doc.add_paragraph()
    run = section_title.add_run("Historical items")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True

    historical_items = [
        ("H1", "History of problems with violence"),
        ("H2", "History of problems with other anti-social behaviour"),
        ("H3", "History of problems with relationships"),
        ("H4", "History of problems with employment"),
        ("H5", "History of Problems with Substance Use"),
        ("H6", "History of Problems with Major Mental Disorder"),
        ("H7", "History of Problems with Personality Disorder"),
        ("H8", "History of Problems with Traumatic Experiences"),
        ("H9", "History of Problems with Violent Attitudes"),
        ("H10", "History of Problems with Treatment or Supervision Response"),
    ]

    for code, title in historical_items:
        add_item_header(doc, code, title)

        # Add subsections for specific items
        if code == "H6":
            add_subsection(doc, "Psychotic Disorders:")
            add_content_placeholder(doc)
            add_subsection(doc, "Major Mood Disorders:")
            add_content_placeholder(doc)
            add_subsection(doc, "Other Major Mental Disorders:")
            add_content_placeholder(doc)
        else:
            add_content_placeholder(doc)

        add_presence_relevance_table(doc)

    # ===== CLINICAL ITEMS =====
    doc.add_page_break()
    section_title = doc.add_paragraph()
    run = section_title.add_run("Clinical items")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True

    clinical_items = [
        ("C1", "Recent Problems with Insight", ["Insight into Violence Risk:", "Insight into Need for Treatment:"]),
        ("C2", "Recent Problems with Violent Ideation or Intent", []),
        ("C3", "Recent Problems with Symptoms of Major Mental Disorder", ["Symptoms of Psychotic Disorders:", "Symptoms of Major Mood Disorder:"]),
        ("C4", "Recent Problems with Instability", ["Affective Instability:", "Behavioural Instability:", "Cognitive Instability:"]),
        ("C5", "Recent Problems with Treatment or Supervision Response", ["Problems with Compliance:"]),
    ]

    for code, title, subsections in clinical_items:
        add_item_header(doc, code, title)

        if subsections:
            for sub in subsections:
                add_subsection(doc, sub)
                add_content_placeholder(doc)
        else:
            add_content_placeholder(doc)

        add_presence_relevance_table(doc)

    # ===== RISK MANAGEMENT ITEMS =====
    doc.add_page_break()
    section_title = doc.add_paragraph()
    run = section_title.add_run("Risk Management items")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True

    risk_items = [
        ("R1", "Future Problems with Professional Services and Plans"),
        ("R2", "Future Problems with Living Situation"),
        ("R3", "Future Problems with Personal Support"),
        ("R4", "Future Problems with Treatment or Supervision Response"),
        ("R5", "Future Problems with Stress or Coping"),
    ]

    for code, title in risk_items:
        add_item_header(doc, code, title)

        # Hospital and Community subsections
        add_subsection(doc, "Hospital:")
        add_content_placeholder(doc)
        add_subsection(doc, "Community:")
        add_content_placeholder(doc)

        add_presence_relevance_table(doc)

    # ===== VIOLENCE RISK FORMULATION =====
    doc.add_page_break()
    section_title = doc.add_paragraph()
    run = section_title.add_run("Violence risk formulation")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(14)
    run.bold = True
    run.underline = True

    # Case Formulation
    para = doc.add_paragraph()
    run = para.add_run("Case Formulation:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # Scenarios
    para = doc.add_paragraph()
    run = para.add_run("Scenarios (what kind of violence is likely to be committed, victims and likely motivation):")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # Seriousness
    para = doc.add_paragraph()
    run = para.add_run("Seriousness (psychological/physical harm to victims, could this escalate to a serious or life-threatening level?):")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # Imminence
    para = doc.add_paragraph()
    run = para.add_run("Imminence (how soon could the individual engage in violence? What are the warning signs that risk is increasing or imminent?):")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # Frequency/Likelihood
    para = doc.add_paragraph()
    run = para.add_run("Frequency/Likelihood (how often might this violence occur? Is the risk chronic or acute? How likely is it that this type of violence will occur?):")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # ===== PROPOSED MANAGEMENT STRATEGIES =====
    doc.add_paragraph()
    section_title = doc.add_paragraph()
    run = section_title.add_run("Proposed management strategies")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(12)
    run.bold = True
    run.underline = True

    # Risk-enhancing factors
    para = doc.add_paragraph()
    run = para.add_run("Risk-enhancing factors (factors that may increase risk):")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # Protective factors
    para = doc.add_paragraph()
    run = para.add_run("Protective factors:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # Monitoring
    para = doc.add_paragraph()
    run = para.add_run("Monitoring:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    add_content_placeholder(doc)

    # ===== TREATMENT / RECOMMENDATIONS =====
    doc.add_paragraph()
    section_title = doc.add_paragraph()
    run = section_title.add_run("Treatment/ Recommendations")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(12)
    run.bold = True
    run.underline = True

    para = doc.add_paragraph()
    run = para.add_run("The following forms of treatment are advised to reduce risk as much as possible:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    add_content_placeholder(doc)

    # Supervision
    para = doc.add_paragraph()
    run = para.add_run("Supervision")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run("The following forms of supervision are advised to reduce risk as much as possible:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    add_content_placeholder(doc)

    # Victim safety planning
    para = doc.add_paragraph()
    run = para.add_run("Victim safety planning")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run("The following strategies are likely to enhance victim safety:")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    add_content_placeholder(doc)

    # ===== SIGNATURE =====
    doc.add_paragraph()
    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run("[Author Name]")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    para = doc.add_paragraph()
    run = para.add_run("[Title]")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)

    para = doc.add_paragraph()
    run = para.add_run("[Date]")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)

    return doc

if __name__ == "__main__":
    doc = create_hcr20_template()
    output_path = "/Users/avie/Desktop/HCR-20_Blank_Template.docx"
    doc.save(output_path)
    print(f"Template saved to: {output_path}")
