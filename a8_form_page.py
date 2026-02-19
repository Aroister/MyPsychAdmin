# ================================================================
#  A8 FORM PAGE — Section 3 Medical Recommendation for Treatment
#  Mental Health Act 1983 - Form A8 Regulation 4(1)(d)(ii)
#  Single practitioner (unlike A7 which is joint)
#  CARD/POPUP LAYOUT with ResizableSection
# ================================================================

from __future__ import annotations
import re
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QGroupBox, QToolButton, QRadioButton,
    QButtonGroup, QComboBox, QSpinBox, QCompleter, QStyleFactory,
    QSlider, QStackedWidget, QSplitter
)

from background_history_popup import ResizableSection
from utils.resource_path import resource_path

# ICD-10 data - curated list matching iOS app
try:
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []

from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# NO-WHEEL SLIDER (prevents scroll from changing value)
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()


class NoWheelSpinBox(QSpinBox):
    def wheelEvent(self, event):
        event.ignore()


class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# A8 CARD WIDGET - Fixed header, scrollable content
# ================================================================
class A8CardWidget(QFrame):
    """Card with fixed header and scrollable content area."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("a8Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#a8Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#a8Card:hover {
                border-color: #7c3aed;
                background: #faf5ff;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QLabel#cardTitle {
                font-size: 20px;
                font-weight: 600;
                color: #7c3aed;
                padding-bottom: 4px;
            }
            QFrame#divider {
                background: #e5e7eb;
                height: 1px;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 12, 18, 12)
        layout.setSpacing(0)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setObjectName("cardTitle")
        self.title_label.setFixedHeight(24)
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        # Divider
        divider = QFrame()
        divider.setObjectName("divider")
        divider.setFixedHeight(1)
        layout.addWidget(divider)

        # Scrollable content
        self.content_scroll = QScrollArea()
        self.content_scroll.setWidgetResizable(True)
        self.content_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        self.content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.content_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        content_widget = QWidget()
        content_widget.setStyleSheet("background: transparent;")
        self.content_layout = QVBoxLayout(content_widget)
        self.content_layout.setContentsMargins(0, 8, 0, 4)
        self.content_layout.setSpacing(0)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 20px;
                color: #374151;
                padding: 4px;
                background: transparent;
                border: none;
            }
        """)
        self.content.setFrameShape(QFrame.Shape.NoFrame)
        self.content_layout.addWidget(self.content)
        self.content_layout.addStretch()

        self.content_scroll.setWidget(content_widget)
        layout.addWidget(self.content_scroll, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# TOOLBAR
# ================================================================
class A8Toolbar(QWidget):
    """Toolbar for the A8 Form Page."""

    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            A8Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN A8 FORM PAGE - Card/Popup Layout
# ================================================================
class A8FormPage(QWidget):
    """Page for completing MHA Form A8 - Section 3 Single Medical Recommendation."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean",
        "Asian",
        "Caucasian",
        "Middle Eastern",
        "Mixed Race",
        "Not specified",
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()
        self.cards = {}

        self._setup_ui()
        self._prefill_practitioner()
        self._connect_clinical_live_preview()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {
            "full_name": details[1] or "",
            "email": details[7] or "",
        }

    def _prefill_practitioner(self):
        if self._my_details.get("full_name"):
            self.prac_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.prac_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #7c3aed; border-bottom: 1px solid #6d28d9;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Form A8 — Section 3 Medical Recommendation (Single)")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)
        from shared_data_store import get_shared_store
        self._shared_store = get_shared_store()
        self._shared_store.uploaded_documents_changed.connect(self._refresh_upload_menu)
        self._refresh_upload_menu(self._shared_store.get_uploaded_documents())

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area - cards on left, popup on right (QSplitter for drag resize)
        content = QSplitter(Qt.Orientation.Horizontal)
        content.setHandleWidth(6)
        content.setStyleSheet("""
            QSplitter { background: #f9fafb; }
            QSplitter::handle { background: #d1d5db; }
            QSplitter::handle:hover { background: #6BAF8D; }
        """)

        # Left: Cards column (scrollable)
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(16, 16, 16, 16)
        self.cards_layout.setSpacing(8)

        # Create cards with ResizableSection
        self._create_patient_card()
        self._create_practitioner_card()
        self._create_clinical_card()
        self._create_signature_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        content.addWidget(cards_scroll)

        # Right: Popup panel (resizable via splitter)
        self.popup_stack = QStackedWidget()
        self.popup_stack.setMinimumWidth(400)
        self.popup_stack.setMaximumWidth(650)
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: none;
                border-radius: 12px;
            }
        """)

        # Create popup panels
        self._create_patient_popup()
        self._create_practitioner_popup()
        self._create_clinical_popup()
        self._create_signature_popup()

        content.addWidget(self.popup_stack)
        content.setSizes([500, 520])
        main_layout.addWidget(content, 1)

        # Initialize cards with default date values
        self._update_practitioner_card()
        self._update_signature_card()

        # Show first popup by default
        self._on_card_clicked("patient")

    def _on_card_clicked(self, key: str):
        """Handle card click - show corresponding popup."""
        index_map = {"patient": 0, "practitioner": 1, "clinical": 2, "signatures": 3}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_patient_card(self):
        section = ResizableSection()
        section.set_content_height(130)
        section._min_height = 80
        section._max_height = 200

        card = A8CardWidget("Patient", "patient")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["patient"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_practitioner_card(self):
        section = ResizableSection()
        section.set_content_height(130)
        section._min_height = 80
        section._max_height = 200

        card = A8CardWidget("Practitioner", "practitioner")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["practitioner"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_clinical_card(self):
        section = ResizableSection()
        section.set_content_height(170)
        section._min_height = 100
        section._max_height = 300

        card = A8CardWidget("Clinical Reasons", "clinical")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["clinical"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signature_card(self):
        section = ResizableSection()
        section.set_content_height(110)
        section._min_height = 80
        section._max_height = 150

        card = A8CardWidget("Treatment & Signature", "signatures")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signatures"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # HELPER METHODS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 17px;
            }
            QLineEdit:focus { border-color: #7c3aed; }
        """)
        return edit

    def _create_date_edit(self) -> NoWheelDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 10px 12px;
                font-size: 17px;
            }
            QDateEdit::drop-down { border: none; width: 24px; }
        """)
        return date_edit

    def _create_styled_frame(self, color: str) -> QFrame:
        """Create a styled frame with colored background."""
        colors = {
            "purple": ("#faf5ff", "#e9d5ff"),
            "green": ("#f0fdf4", "#bbf7d0"),
            "blue": ("#eff6ff", "#bfdbfe"),
            "yellow": ("#fefce8", "#fef08a"),
            "red": ("#fef2f2", "#fecaca"),
        }
        bg, border = colors.get(color, ("#f9fafb", "#e5e7eb"))
        frame = QFrame()
        frame.setStyleSheet(f"QFrame {{ background: {bg}; border: none; border-radius: 8px; }}")
        return frame

    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_patient_popup(self):
        """Create patient details popup (Demographics + Patient combined)."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        # Header
        header = QLabel("Patient Details")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Patient frame (purple)
        patient_frame = self._create_styled_frame("purple")
        pf_layout = QVBoxLayout(patient_frame)
        pf_layout.setContentsMargins(16, 12, 16, 12)
        pf_layout.setSpacing(12)

        # Name
        name_lbl = QLabel("Full Name:")
        name_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #7c3aed;")
        pf_layout.addWidget(name_lbl)
        self.patient_name = self._create_line_edit("Patient's full name")
        pf_layout.addWidget(self.patient_name)

        # Address
        addr_lbl = QLabel("Address:")
        addr_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #7c3aed;")
        pf_layout.addWidget(addr_lbl)
        self.patient_address = self._create_line_edit("Patient's address")
        pf_layout.addWidget(self.patient_address)

        layout.addWidget(patient_frame)

        # Demographics frame (green)
        demo_frame = self._create_styled_frame("green")
        df_layout = QVBoxLayout(demo_frame)
        df_layout.setContentsMargins(16, 12, 16, 12)
        df_layout.setSpacing(12)

        demo_header = QLabel("Demographics")
        demo_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #166534;")
        df_layout.addWidget(demo_header)

        # Demographics row with Age and Gender
        demo_row = QHBoxLayout()
        demo_row.setSpacing(12)

        # Age
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #166534;")
        demo_row.addWidget(age_lbl)
        self.age_spin = NoWheelSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setValue(0)
        self.age_spin.setStyleSheet("font-size: 17px; padding: 4px;")
        self.age_spin.setFixedWidth(70)
        demo_row.addWidget(self.age_spin)

        demo_row.addSpacing(20)

        # Gender
        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("Other")
        self.gender_group.addButton(self.gender_male, 0)
        self.gender_group.addButton(self.gender_female, 1)
        self.gender_group.addButton(self.gender_other, 2)
        demo_row.addWidget(self.gender_male)
        demo_row.addWidget(self.gender_female)
        demo_row.addWidget(self.gender_other)

        demo_row.addStretch()
        df_layout.addLayout(demo_row)

        # Ethnicity row (separate line)
        eth_row = QHBoxLayout()
        eth_row.setSpacing(12)
        eth_lbl = QLabel("Ethnicity:")
        eth_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #166534;")
        eth_row.addWidget(eth_lbl)
        self.ethnicity_combo = NoWheelComboBox()
        self.ethnicity_combo.addItem("Not specified")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setStyleSheet("font-size: 17px; padding: 4px;")
        self.ethnicity_combo.setMinimumWidth(200)
        eth_row.addWidget(self.ethnicity_combo)
        eth_row.addStretch()
        df_layout.addLayout(eth_row)

        layout.addWidget(demo_frame)

        layout.addStretch()

        # Auto-sync to card (age and ethnicity are backend only)
        self.patient_name.textChanged.connect(self._update_patient_card)
        self.patient_address.textChanged.connect(self._update_patient_card)

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_practitioner_popup(self):
        """Create single practitioner popup."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        # Header
        header = QLabel("Medical Practitioner")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Practitioner frame (blue)
        prac_frame = self._create_styled_frame("blue")
        pf_layout = QVBoxLayout(prac_frame)
        pf_layout.setContentsMargins(16, 12, 16, 12)
        pf_layout.setSpacing(12)

        # Name
        name_lbl = QLabel("Full Name:")
        name_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #1e40af;")
        pf_layout.addWidget(name_lbl)
        self.prac_name = self._create_line_edit("Practitioner's full name")
        pf_layout.addWidget(self.prac_name)

        # Address
        addr_lbl = QLabel("Address:")
        addr_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #1e40af;")
        pf_layout.addWidget(addr_lbl)
        self.prac_address = self._create_line_edit("Practitioner's address")
        pf_layout.addWidget(self.prac_address)

        # Email
        email_lbl = QLabel("Email:")
        email_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #1e40af;")
        pf_layout.addWidget(email_lbl)
        self.prac_email = self._create_line_edit("Email address")
        pf_layout.addWidget(self.prac_email)

        # Exam date row
        exam_row = QHBoxLayout()
        exam_lbl = QLabel("Examined on:")
        exam_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #1e40af;")
        exam_row.addWidget(exam_lbl)
        self.prac_exam_date = self._create_date_edit()
        self.prac_exam_date.setFixedWidth(140)
        exam_row.addWidget(self.prac_exam_date)
        exam_row.addStretch()
        pf_layout.addLayout(exam_row)

        # Checkboxes
        self.prac_acquaintance = QCheckBox("Previous acquaintance with patient")
        self.prac_acquaintance.setStyleSheet("font-size: 17px; color: #374151;")
        pf_layout.addWidget(self.prac_acquaintance)

        self.prac_section12 = QCheckBox("Section 12 approved")
        self.prac_section12.setStyleSheet("font-size: 17px; color: #374151; font-weight: 600;")
        pf_layout.addWidget(self.prac_section12)

        layout.addWidget(prac_frame)

        # S12 note
        note = QLabel("NOTE: The practitioner completing this form must be approved under Section 12 of the Act.")
        note.setWordWrap(True)
        note.setStyleSheet("font-size: 19px; color: #dc2626; padding: 8px; background: #fef2f2; border-radius: 4px; font-weight: 500;")
        layout.addWidget(note)

        layout.addStretch()

        # Auto-sync to card
        self.prac_name.textChanged.connect(self._update_practitioner_card)
        self.prac_address.textChanged.connect(self._update_practitioner_card)
        self.prac_exam_date.dateChanged.connect(self._update_practitioner_card)
        self.prac_acquaintance.toggled.connect(self._update_practitioner_card)
        self.prac_section12.toggled.connect(self._update_practitioner_card)

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_clinical_popup(self):
        """Create clinical reasons popup with controls and live preview."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        container.setMaximumWidth(500)
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        # Header
        header = QLabel("Clinical Reasons")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Hidden label for state storage (used in export and state management)
        self.clinical_reasons = QLabel("")
        self.clinical_reasons.hide()

        # Controls section
        controls_frame = QFrame()
        controls_frame.setStyleSheet("QFrame { background: #f9fafb; border: none; border-radius: 8px; }")
        ctrl_layout = QVBoxLayout(controls_frame)
        ctrl_layout.setContentsMargins(12, 12, 12, 12)
        ctrl_layout.setSpacing(12)

        # Mental Disorder (ICD-10)
        md_frame = self._create_styled_frame("green")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(8)

        md_header = QLabel("Mental Disorder (ICD-10)")
        md_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        # Primary diagnosis dropdown with grouped items
        self.dx_primary = NoWheelComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select primary diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        md_layout.addWidget(self.dx_primary)

        # Secondary diagnosis dropdown
        self.dx_secondary = NoWheelComboBox()
        self.dx_secondary.setEditable(True)
        self.dx_secondary.lineEdit().setPlaceholderText("Secondary diagnosis (optional)...")
        self.dx_secondary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_secondary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_secondary.addItem(f"── {group_name} ──", None)
            idx = self.dx_secondary.count() - 1
            self.dx_secondary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_secondary.addItem(dx, dx)
        completer2 = QCompleter(ICD10_FLAT, self.dx_secondary)
        completer2.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer2.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_secondary.setCompleter(completer2)
        self.dx_secondary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        md_layout.addWidget(self.dx_secondary)

        ctrl_layout.addWidget(md_frame)

        # Legal Criteria
        lc_frame = self._create_styled_frame("blue")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(6)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature with sub-options
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.relapsing_cb.setStyleSheet("font-size: 20px; color: #6b7280;")
        nature_opt_layout.addWidget(self.relapsing_cb)

        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.treatment_resistant_cb.setStyleSheet("font-size: 20px; color: #6b7280;")
        nature_opt_layout.addWidget(self.treatment_resistant_cb)

        self.chronic_cb = QCheckBox("Chronic and enduring")
        self.chronic_cb.setStyleSheet("font-size: 20px; color: #6b7280;")
        nature_opt_layout.addWidget(self.chronic_cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree with slider
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setTickPosition(QSlider.TickPosition.TicksBelow)
        self.degree_slider.setTickInterval(1)
        self.degree_slider.setFixedWidth(120)
        slider_row.addWidget(self.degree_slider)

        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 20px; color: #374151; font-weight: 500;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 20px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health with sub-options
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 20px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 20px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)

        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.poor_compliance_cb.setStyleSheet("font-size: 20px; color: #9ca3af;")
        mh_opt_layout.addWidget(self.poor_compliance_cb)

        self.limited_insight_cb = QCheckBox("Limited insight")
        self.limited_insight_cb.setStyleSheet("font-size: 20px; color: #9ca3af;")
        mh_opt_layout.addWidget(self.limited_insight_cb)

        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 20px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 20px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety with sub-options
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 20px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # Self section
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 20px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 20px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)

        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_neglect.setStyleSheet("font-size: 20px; color: #9ca3af;")
        self_opt_layout.addWidget(self.self_hist_neglect)

        self.self_hist_risky = QCheckBox("Self placement in risky situations")
        self.self_hist_risky.setStyleSheet("font-size: 20px; color: #9ca3af;")
        self_opt_layout.addWidget(self.self_hist_risky)

        self.self_hist_harm = QCheckBox("Self harm")
        self.self_hist_harm.setStyleSheet("font-size: 20px; color: #9ca3af;")
        self_opt_layout.addWidget(self.self_hist_harm)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 20px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        self_opt_layout.addWidget(self_curr_lbl)

        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_neglect.setStyleSheet("font-size: 20px; color: #9ca3af;")
        self_opt_layout.addWidget(self.self_curr_neglect)

        self.self_curr_risky = QCheckBox("Self placement in risky situations")
        self.self_curr_risky.setStyleSheet("font-size: 20px; color: #9ca3af;")
        self_opt_layout.addWidget(self.self_curr_risky)

        self.self_curr_harm = QCheckBox("Self harm")
        self.self_curr_harm.setStyleSheet("font-size: 20px; color: #9ca3af;")
        self_opt_layout.addWidget(self.self_curr_harm)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others section
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 20px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 20px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)

        self.others_hist_violence = QCheckBox("Violence to others")
        self.others_hist_violence.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_hist_violence)

        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_verbal.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_hist_verbal)

        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_sexual.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_hist_sexual)

        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_stalking.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_hist_stalking)

        self.others_hist_arson = QCheckBox("Arson")
        self.others_hist_arson.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_hist_arson)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 20px; color: #9ca3af; font-weight: 600; margin-top: 4px;")
        others_opt_layout.addWidget(others_curr_lbl)

        self.others_curr_violence = QCheckBox("Violence to others")
        self.others_curr_violence.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_curr_violence)

        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_verbal.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_curr_verbal)

        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_sexual.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_curr_sexual)

        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_stalking.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_curr_stalking)

        self.others_curr_arson = QCheckBox("Arson")
        self.others_curr_arson.setStyleSheet("font-size: 20px; color: #9ca3af;")
        others_opt_layout.addWidget(self.others_curr_arson)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        ctrl_layout.addWidget(lc_frame)

        # Informal Not Indicated
        inf_frame = self._create_styled_frame("red")
        inf_layout = QVBoxLayout(inf_frame)
        inf_layout.setContentsMargins(12, 10, 12, 10)
        inf_layout.setSpacing(4)

        inf_header = QLabel("Informal Not Indicated")
        inf_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #991b1b;")
        inf_layout.addWidget(inf_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed")
        self.tried_failed_cb.setStyleSheet("font-size: 20px; color: #374151;")
        inf_layout.addWidget(self.tried_failed_cb)

        self.insight_cb = QCheckBox("Lack of Insight")
        self.insight_cb.setStyleSheet("font-size: 20px; color: #374151;")
        inf_layout.addWidget(self.insight_cb)

        self.compliance_cb = QCheckBox("Compliance Issues")
        self.compliance_cb.setStyleSheet("font-size: 20px; color: #374151;")
        inf_layout.addWidget(self.compliance_cb)

        self.supervision_cb = QCheckBox("Needs MHA Supervision")
        self.supervision_cb.setStyleSheet("font-size: 20px; color: #374151;")
        inf_layout.addWidget(self.supervision_cb)

        ctrl_layout.addWidget(inf_frame)

        layout.addWidget(controls_frame)
        layout.addStretch()

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_signature_popup(self):
        """Create treatment & signature popup."""
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(16)

        # Header
        header = QLabel("Treatment & Signature")
        header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Hospital frame (yellow)
        hosp_frame = self._create_styled_frame("yellow")
        hf_layout = QVBoxLayout(hosp_frame)
        hf_layout.setContentsMargins(16, 12, 16, 12)
        hf_layout.setSpacing(12)

        hosp_header = QLabel("Appropriate Treatment")
        hosp_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #854d0e;")
        hf_layout.addWidget(hosp_header)

        info = QLabel("Hospital where appropriate treatment is available:")
        info.setWordWrap(True)
        info.setStyleSheet("font-size: 20px; color: #6b7280;")
        hf_layout.addWidget(info)

        self.hospital_treatment = self._create_line_edit("Hospital name(s)")
        hf_layout.addWidget(self.hospital_treatment)

        layout.addWidget(hosp_frame)

        # Signature frame (purple)
        sig_frame = self._create_styled_frame("purple")
        sf_layout = QVBoxLayout(sig_frame)
        sf_layout.setContentsMargins(16, 12, 16, 12)
        sf_layout.setSpacing(12)

        sig_header = QLabel("Signature")
        sig_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #7c3aed;")
        sf_layout.addWidget(sig_header)

        sig_row = QHBoxLayout()
        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)

        self.sig_date = self._create_date_edit()
        self.sig_date.setFixedWidth(140)
        sig_row.addWidget(self.sig_date)
        sig_row.addStretch()
        sf_layout.addLayout(sig_row)

        layout.addWidget(sig_frame)

        layout.addStretch()

        # Auto-sync to card
        self.hospital_treatment.textChanged.connect(self._update_signature_card)
        self.sig_date.dateChanged.connect(self._update_signature_card)

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_patient_card(self):
        parts = []
        if self.patient_name.text():
            parts.append(self.patient_name.text())
        if self.patient_address.text():
            parts.append(self.patient_address.text())
        self.cards["patient"].set_content_text("\n".join(parts) if parts else "No patient details entered")

    def _update_practitioner_card(self):
        parts = []
        if self.prac_name.text():
            parts.append(self.prac_name.text())
        if self.prac_address.text():
            parts.append(self.prac_address.text())

        details = []
        exam_date = self.prac_exam_date.date().toString("dd MMM yyyy")
        parts.append(exam_date)

        self.cards["practitioner"].set_content_text("\n".join(parts) if parts else "No practitioner details entered")

    def _update_clinical_card(self):
        text = self._generate_clinical_text()
        self.cards["clinical"].set_content_text(text if text else "No clinical reasons entered")

    def _update_signature_card(self):
        parts = []
        if self.hospital_treatment.text():
            parts.append(self.hospital_treatment.text())
        sig_date = self.sig_date.date().toString("dd MMM yyyy")
        parts.append(sig_date)

        self.cards["signatures"].set_content_text("\n".join(parts))

    # ----------------------------------------------------------------
    # LIVE PREVIEW
    # ----------------------------------------------------------------
    def _connect_clinical_live_preview(self):
        """Connect all clinical controls to live preview update."""
        # ICD-10 text fields
        self.dx_primary.currentTextChanged.connect(self._update_clinical_preview)
        self.dx_secondary.currentTextChanged.connect(self._update_clinical_preview)

        # Main checkboxes
        for cb in [self.nature_cb, self.degree_cb, self.health_cb, self.safety_cb,
                   self.tried_failed_cb, self.insight_cb, self.compliance_cb, self.supervision_cb]:
            cb.toggled.connect(self._update_clinical_preview)

        # Nature sub-options
        for cb in [self.relapsing_cb, self.treatment_resistant_cb, self.chronic_cb]:
            cb.toggled.connect(self._update_clinical_preview)

        # Degree
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)
        self.degree_details.textChanged.connect(self._update_clinical_preview)

        # Health sub-options
        for cb in [self.mental_health_cb, self.physical_health_cb, self.poor_compliance_cb, self.limited_insight_cb]:
            cb.toggled.connect(self._update_clinical_preview)
        self.physical_health_details.textChanged.connect(self._update_clinical_preview)

        # Safety sub-options
        for cb in [self.self_harm_cb, self.others_cb,
                   self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm,
                   self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm,
                   self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual,
                   self.others_hist_stalking, self.others_hist_arson,
                   self.others_curr_violence, self.others_curr_verbal, self.others_curr_sexual,
                   self.others_curr_stalking, self.others_curr_arson]:
            cb.toggled.connect(self._update_clinical_preview)

        # Patient name for clinical text
        self.patient_name.textChanged.connect(self._update_clinical_preview)

    def _update_clinical_preview(self):
        """Update the clinical reasons and auto-sync to card."""
        text = self._generate_clinical_text()
        self.clinical_reasons.setText(text)
        self.cards["clinical"].set_content_text(text if text else "No clinical reasons entered")

    # ----------------------------------------------------------------
    # CONTROL TOGGLE HANDLERS
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)
        self._update_clinical_preview()

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()
        self._update_clinical_preview()

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))
        self._update_clinical_preview()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)
        self._update_clinical_preview()

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)
        self._update_clinical_preview()

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()
        self._update_clinical_preview()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)
        self._update_clinical_preview()

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            self.self_hist_neglect.setChecked(False)
            self.self_hist_risky.setChecked(False)
            self.self_hist_harm.setChecked(False)
            self.self_curr_neglect.setChecked(False)
            self.self_curr_risky.setChecked(False)
            self.self_curr_harm.setChecked(False)
        self._update_clinical_preview()

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            self.others_hist_violence.setChecked(False)
            self.others_hist_verbal.setChecked(False)
            self.others_hist_sexual.setChecked(False)
            self.others_hist_stalking.setChecked(False)
            self.others_hist_arson.setChecked(False)
            self.others_curr_violence.setChecked(False)
            self.others_curr_verbal.setChecked(False)
            self.others_curr_sexual.setChecked(False)
            self.others_curr_stalking.setChecked(False)
            self.others_curr_arson.setChecked(False)
        self._update_clinical_preview()

    # ----------------------------------------------------------------
    # CLINICAL TEXT GENERATION (uses "I" for single practitioner)
    # ----------------------------------------------------------------
    def _generate_clinical_text(self) -> str:
        """Generate clinical reasons text from form selections (single practitioner - uses 'I')."""
        p = self._get_pronouns()

        # === PARAGRAPH 1: Demographics, Diagnosis, Nature/Degree ===
        para1_parts = []

        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")

        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Select...", "Not specified"):
            eth_simple = ethnicity.replace(" British", "").replace(" Other", "")
            opening_parts.append(eth_simple)

        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        if self.dx_primary.currentText().strip():
            diagnoses.append(self.dx_primary.currentText().strip())
        if self.dx_secondary.currentText().strip():
            diagnoses.append(self.dx_secondary.currentText().strip())

        if diagnoses:
            if opening_parts:
                demo_str = " ".join(opening_parts)
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} is a {demo_str} who suffers from {joined} which are mental disorders as defined by the Mental Health Act.")
            else:
                if len(diagnoses) == 1:
                    para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
                else:
                    joined = ", ".join(diagnoses[:-1]) + f" and {diagnoses[-1]}"
                    para1_parts.append(f"{name_display} suffers from {joined} which are mental disorders as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree to warrant detention for treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature to warrant detention for treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree to warrant detention for treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    nature_str = " and ".join(nature_types)
                    para1_parts.append(f"The nature of the illness is {nature_str}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if "disorder" in dx_name:
                    dx_name = dx_name[:dx_name.index("disorder") + len("disorder")]
                if dx_name.startswith("schizophrenia"):
                    dx_name = "schizophrenia"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name} including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms of {dx_name}.")

        elif opening_parts:
            demo_str = " ".join(opening_parts)
            para1_parts.append(f"{name_display} is a {demo_str}.")

        # === PARAGRAPH 2: Necessity (Health + Safety) - uses "I" ===
        para2_parts = []

        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("safety of others")

        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"Detention for treatment is necessary due to risks to {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")

        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_reasons = []
            if self.poor_compliance_cb.isChecked():
                mh_reasons.append("non compliance")
            if self.limited_insight_cb.isChecked():
                mh_reasons.append("limited insight")

            if mh_reasons:
                reasons_str = "/".join(mh_reasons)
                para2_parts.append(f"Regarding health I would be concerned about {p['pos_l']} mental health deteriorating due to {reasons_str}.")
            else:
                para2_parts.append(f"Regarding health I would be concerned about {p['pos_l']} mental health deteriorating.")

        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health: {details}.")
            else:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health.")

        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            if self.gender_male.isChecked():
                reflexive = "himself"
            elif self.gender_female.isChecked():
                reflexive = "herself"
            else:
                reflexive = "themselves"

            risk_types = [
                ("self neglect", self.self_hist_neglect.isChecked(), self.self_curr_neglect.isChecked()),
                (f"placing of {reflexive} in risky situations", self.self_hist_risky.isChecked(), self.self_curr_risky.isChecked()),
                ("self harm", self.self_hist_harm.isChecked(), self.self_curr_harm.isChecked()),
            ]

            both_items = []
            hist_only = []
            curr_only = []

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            self_text = f"With respect to {p['pos_l']} own safety I am concerned about"
            parts = []
            if both_items:
                if len(both_items) == 1:
                    parts.append(f"historical and current {both_items[0]}")
                else:
                    parts.append(f"historical and current {', '.join(both_items[:-1])}, and of {both_items[-1]}")
            if hist_only:
                if len(hist_only) == 1:
                    parts.append(f"historical {hist_only[0]}")
                else:
                    parts.append(f"historical {', '.join(hist_only[:-1])}, and of {hist_only[-1]}")
            if curr_only:
                if len(curr_only) == 1:
                    parts.append(f"current {curr_only[0]}")
                else:
                    parts.append(f"current {', '.join(curr_only[:-1])}, and of {curr_only[-1]}")

            if parts:
                self_text += " " + ", and ".join(parts) + "."
            else:
                self_text += f" {p['pos_l']} safety."

            para2_parts.append(self_text)

        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            risk_types = [
                ("violence to others", self.others_hist_violence.isChecked(), self.others_curr_violence.isChecked()),
                ("verbal aggression", self.others_hist_verbal.isChecked(), self.others_curr_verbal.isChecked()),
                ("sexual violence", self.others_hist_sexual.isChecked(), self.others_curr_sexual.isChecked()),
                ("stalking", self.others_hist_stalking.isChecked(), self.others_curr_stalking.isChecked()),
                ("arson", self.others_hist_arson.isChecked(), self.others_curr_arson.isChecked()),
            ]

            both_items = []
            hist_only = []
            curr_only = []

            for risk_name, is_hist, is_curr in risk_types:
                if is_hist and is_curr:
                    both_items.append(risk_name)
                elif is_hist:
                    hist_only.append(risk_name)
                elif is_curr:
                    curr_only.append(risk_name)

            others_text = "With respect to risk to others I am concerned about the risk of"
            parts = []
            if both_items:
                parts.append(f"historical and current {', '.join(both_items)}")
            if hist_only:
                parts.append(f"historical {', '.join(hist_only)}")
            if curr_only:
                parts.append(f"current {', '.join(curr_only)}")

            if parts:
                others_text += " " + " and of ".join(parts) + "."
            else:
                others_text += f" {p['pos_l']} potential to cause harm."

            para2_parts.append(others_text)

        # === PARAGRAPH 3: Informal not indicated - uses "I" ===
        para3_parts = []

        if self.tried_failed_cb.isChecked():
            para3_parts.append("Previous attempts at informal admissions have not been successful and I would likewise be concerned about this recurring in this instance hence I do not believe informal admission currently would be appropriate.")

        if self.insight_cb.isChecked():
            para3_parts.append(f"{p['pos']} lack of insight is a significant concern and should {p['subj_l']} be discharged from section, I believe this would significantly impair {p['pos_l']} compliance if informal.")

        if self.compliance_cb.isChecked():
            if para3_parts:
                para3_parts.append(f"Compliance with treatment has also been a significant issue and I do not believe {p['subj_l']} would comply if informal.")
            else:
                para3_parts.append(f"Compliance with treatment has been a significant issue and I do not believe {p['subj_l']} would comply if informal.")

        if self.supervision_cb.isChecked():
            name = patient_name if patient_name else "the patient"
            para3_parts.append(f"I believe {name} needs careful community monitoring under the supervision afforded by the mental health act and I do not believe such supervision would be complied with should {p['subj_l']} remain in the community informally.")

        paragraphs = []
        if para1_parts:
            paragraphs.append(" ".join(para1_parts))
        if para2_parts:
            paragraphs.append(" ".join(para2_parts))
        if para3_parts:
            paragraphs.append(" ".join(para3_parts))

        return "\n\n".join(paragraphs)

    # ----------------------------------------------------------------
    # EDITOR FOCUS TRACKING
    # ----------------------------------------------------------------
    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    # ----------------------------------------------------------------
    # ACTIONS
    # ----------------------------------------------------------------
    def _go_back(self):
        self.go_back.emit()

    def _clear_form(self):
        reply = QMessageBox.question(
            self, "Clear Form", "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.age_spin.setValue(0)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.ethnicity_combo.setCurrentIndex(0)
            self.patient_name.clear()
            self.patient_address.clear()
            self.prac_name.clear()
            self.prac_address.clear()
            self.prac_email.clear()
            self.prac_exam_date.setDate(QDate.currentDate())
            self.prac_acquaintance.setChecked(False)
            self.prac_section12.setChecked(False)
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.degree_details.clear()
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.clinical_reasons.setText("")
            self.hospital_treatment.clear()
            self.sig_date.setDate(QDate.currentDate())
            # Reset ALL checkboxes (catches sub-detail checkboxes)
            for cb in self.findChildren(QCheckBox):
                cb.setChecked(False)
            # Update all cards
            self._update_patient_card()
            self._update_practitioner_card()
            self._update_clinical_card()
            self._update_signature_card()
            # Restore my details fields
            self._prefill_practitioner()

    def _refresh_upload_menu(self, docs=None):
        """Rebuild the Uploaded Docs dropdown menu from SharedDataStore."""
        menu = self.toolbar.upload_menu
        menu.clear()
        if docs is None:
            from shared_data_store import get_shared_store
            docs = get_shared_store().get_uploaded_documents()
        if not docs:
            action = menu.addAction("No documents uploaded")
            action.setEnabled(False)
        else:
            for doc in docs:
                path = doc["path"]
                action = menu.addAction(doc["filename"])
                action.triggered.connect(lambda checked=False, p=path: self._import_from_upload(p))

    def _import_from_upload(self, file_path):
        """Process an uploaded file."""
        QMessageBox.information(self, "Import", f"File selected: {file_path}\n\nData extraction coming soon.")

    def _export_docx(self):
        """Export the form to DOCX format with S12 validation."""
        # S12 validation - must be checked for single practitioner form
        if not self.prac_section12.isChecked():
            QMessageBox.warning(
                self, "S12 Approval Required",
                "Cannot export: The practitioner must be Section 12 approved for Form A8."
            )
            self._on_card_clicked("practitioner")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form A8",
            f"Form_A8_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_A8_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form A8 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)

            # Clean ALL paragraphs - remove permission markers and convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                # Remove permission markers (they cause grey appearance)
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                # Convert paragraph-level grey shading to cream
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), 'FFFED5')
                # Convert run-level grey shading to cream
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), 'FFFED5')

            def set_entry_box(para, content=""):
                """Set entry box with bold gold brackets [content]"""
                if not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr = bracket_open._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'FFFED5')
                rPr2.append(shd2)
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr3 = bracket_close._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            def highlight_yellow(para):
                """Rebuild paragraph with cream highlight"""
                text = ''.join(run.text for run in para.runs)
                if not text:
                    text = para.text
                if not text.strip():
                    return
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                    rPr = para.runs[0]._element.get_or_add_rPr()
                    for old_shd in rPr.findall(qn('w:shd')):
                        rPr.remove(old_shd)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)
                else:
                    run = para.add_run(text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    rPr = run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)

            def add_gold_bracket_start(para):
                """Add gold bracket [ at the start of paragraph"""
                text = ''.join(run.text for run in para.runs)
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = '['
                    para.runs[0].font.bold = True
                    para.runs[0].font.color.rgb = BRACKET_COLOR
                else:
                    bracket = para.add_run('[')
                    bracket.font.bold = True
                    bracket.font.color.rgb = BRACKET_COLOR
                content_run = para.add_run(text)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr = content_run._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)

            def add_gold_bracket_end(para):
                """Add gold bracket ] at the end of paragraph"""
                text = ''.join(run.text for run in para.runs)
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                    rPr = para.runs[0]._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)
                else:
                    content_run = para.add_run(text)
                    content_run.font.name = 'Arial'
                    content_run.font.size = Pt(12)
                    rPr = content_run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)
                bracket = para.add_run(']')
                bracket.font.bold = True
                bracket.font.color.rgb = BRACKET_COLOR

            paragraphs = doc.paragraphs

            # Debug: Print paragraph indices
            print("DEBUG A8: Paragraphs 0-50:")
            for i in range(min(50, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)[:60]
                print(f"DEBUG: Para {i}: {txt}")

            # Practitioner entry box with gold brackets (include email)
            prac_parts = []
            if self.prac_name.text().strip():
                prac_parts.append(self.prac_name.text().strip())
            if self.prac_address.text().strip():
                prac_parts.append(self.prac_address.text().strip())
            if self.prac_email.text().strip():
                prac_parts.append(self.prac_email.text().strip())
            if prac_parts:
                set_entry_box(paragraphs[3], ", ".join(prac_parts))
            else:
                set_entry_box(paragraphs[3], "")

            # Patient entry box with gold brackets
            patient_text = self.cards["patient"].get_content().replace("\n", ", ")
            if patient_text.strip() and patient_text != "No patient details entered":
                set_entry_box(paragraphs[5], patient_text)
            else:
                set_entry_box(paragraphs[5], "")

            # Exam date entry box with gold brackets
            exam_date = self.prac_exam_date.date().toString("dd MMMM yyyy")
            set_entry_box(paragraphs[8], exam_date)

            # Acquaintance/Section 12 with gold brackets around group
            # Para 9: *I had previous acquaintance... (gold bracket [ at start)
            # Para 10: *I am approved under section 12... (gold bracket ] at end)
            add_gold_bracket_start(paragraphs[9])
            add_gold_bracket_end(paragraphs[10])
            if not self.prac_acquaintance.isChecked():
                strikethrough_para(paragraphs[9])
            if not self.prac_section12.isChecked():
                strikethrough_para(paragraphs[10])

            # Necessity options - paragraphs 16, 17, 18 in A8 template
            # A8 template has NO (i)(ii)(iii) markers - just the text:
            # Para 16: "for the patient's own health"
            # Para 17: "for the patient's own safety"
            # Para 18: "for the protection of other persons"
            # Format: [for the patient's own health
            #         for the patient's own safety
            #         for the protection of other persons]

            # Para 16: gold [ at start, then highlighted text
            p16 = paragraphs[16]
            p16_text = ''.join(run.text for run in p16.runs)
            for run in p16.runs:
                run.text = ""
            while len(p16.runs) > 1:
                p16._element.remove(p16.runs[-1]._element)
            if p16.runs:
                p16.runs[0].text = '['
                p16.runs[0].font.bold = True
                p16.runs[0].font.color.rgb = BRACKET_COLOR
            else:
                p16_open = p16.add_run('[')
                p16_open.font.bold = True
                p16_open.font.color.rgb = BRACKET_COLOR
            p16_content = p16.add_run(p16_text.strip())
            p16_content.font.name = 'Arial'
            p16_content.font.size = Pt(12)
            rPr16c = p16_content._element.get_or_add_rPr()
            shd16c = OxmlElement('w:shd')
            shd16c.set(qn('w:val'), 'clear')
            shd16c.set(qn('w:color'), 'auto')
            shd16c.set(qn('w:fill'), 'FFFED5')
            rPr16c.append(shd16c)

            # Para 17: just highlighted text (no brackets)
            p17 = paragraphs[17]
            p17_text = ''.join(run.text for run in p17.runs)
            for run in p17.runs:
                run.text = ""
            while len(p17.runs) > 1:
                p17._element.remove(p17.runs[-1]._element)
            if p17.runs:
                p17.runs[0].text = p17_text.strip()
                p17.runs[0].font.name = 'Arial'
                p17.runs[0].font.size = Pt(12)
                rPr17 = p17.runs[0]._element.get_or_add_rPr()
                shd17 = OxmlElement('w:shd')
                shd17.set(qn('w:val'), 'clear')
                shd17.set(qn('w:color'), 'auto')
                shd17.set(qn('w:fill'), 'FFFED5')
                rPr17.append(shd17)
            else:
                p17_content = p17.add_run(p17_text.strip())
                p17_content.font.name = 'Arial'
                p17_content.font.size = Pt(12)
                rPr17c = p17_content._element.get_or_add_rPr()
                shd17c = OxmlElement('w:shd')
                shd17c.set(qn('w:val'), 'clear')
                shd17c.set(qn('w:color'), 'auto')
                shd17c.set(qn('w:fill'), 'FFFED5')
                rPr17c.append(shd17c)

            # Para 18: highlighted text, then gold ] at end
            p18 = paragraphs[18]
            p18_text = ''.join(run.text for run in p18.runs)
            for run in p18.runs:
                run.text = ""
            while len(p18.runs) > 1:
                p18._element.remove(p18.runs[-1]._element)
            if p18.runs:
                p18.runs[0].text = p18_text.strip()
                p18.runs[0].font.name = 'Arial'
                p18.runs[0].font.size = Pt(12)
                rPr18 = p18.runs[0]._element.get_or_add_rPr()
                shd18 = OxmlElement('w:shd')
                shd18.set(qn('w:val'), 'clear')
                shd18.set(qn('w:color'), 'auto')
                shd18.set(qn('w:fill'), 'FFFED5')
                rPr18.append(shd18)
            else:
                p18_content = p18.add_run(p18_text.strip())
                p18_content.font.name = 'Arial'
                p18_content.font.size = Pt(12)
                rPr18c = p18_content._element.get_or_add_rPr()
                shd18c = OxmlElement('w:shd')
                shd18c.set(qn('w:val'), 'clear')
                shd18c.set(qn('w:color'), 'auto')
                shd18c.set(qn('w:fill'), 'FFFED5')
                rPr18c.append(shd18c)
            p18_close = p18.add_run(']')
            p18_close.font.bold = True
            p18_close.font.color.rgb = BRACKET_COLOR

            # Strikethrough unselected detention reasons
            if not self.health_cb.isChecked():
                strikethrough_para(paragraphs[16])
            if not (self.safety_cb.isChecked() and self.self_harm_cb.isChecked()):
                strikethrough_para(paragraphs[17])
            if not (self.safety_cb.isChecked() and self.others_cb.isChecked()):
                strikethrough_para(paragraphs[18])

            # "that this patient should receive treatment in hospital," - NO highlight (matches A4)
            # "(c) such treatment cannot be provided..." - NO highlight (matches A4)

            # Clinical reasons entry box with gold brackets
            clinical_text = self.cards["clinical"].get_content()
            if clinical_text and clinical_text != "No clinical reasons entered":
                # Find clinical reasons paragraph (usually around para 24)
                for i in range(20, min(35, len(paragraphs))):
                    txt = ''.join(run.text for run in paragraphs[i].runs)
                    # Look for an empty or entry box paragraph after "because"
                    if not txt.strip() or txt.strip() in ['[', ']', '[ ]']:
                        set_entry_box(paragraphs[i], clinical_text)
                        break
                else:
                    set_entry_box(paragraphs[24], clinical_text)
            else:
                set_entry_box(paragraphs[24], "")

            # "[If you need to continue...]" with gold checkbox
            for i in range(20, min(40, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "If you need to continue" in txt:
                    para_cont = paragraphs[i]
                    for run in para_cont.runs:
                        run.text = ""
                    while len(para_cont.runs) > 1:
                        para_cont._element.remove(para_cont.runs[-1]._element)
                    if para_cont.runs:
                        para_cont.runs[0].text = '[If you need to continue on a separate sheet please indicate here '
                        para_cont.runs[0].font.name = 'Arial'
                        para_cont.runs[0].font.size = Pt(12)
                    else:
                        r = para_cont.add_run('[If you need to continue on a separate sheet please indicate here ')
                        r.font.name = 'Arial'
                        r.font.size = Pt(12)
                    cb_open = para_cont.add_run('[')
                    cb_open.font.bold = True
                    cb_open.font.color.rgb = BRACKET_COLOR
                    rPr_cbo = cb_open._element.get_or_add_rPr()
                    shd_cbo = OxmlElement('w:shd')
                    shd_cbo.set(qn('w:val'), 'clear')
                    shd_cbo.set(qn('w:color'), 'auto')
                    shd_cbo.set(qn('w:fill'), 'FFFED5')
                    rPr_cbo.append(shd_cbo)
                    cb_space = para_cont.add_run(' ')
                    rPr_cbs = cb_space._element.get_or_add_rPr()
                    shd_cbs = OxmlElement('w:shd')
                    shd_cbs.set(qn('w:val'), 'clear')
                    shd_cbs.set(qn('w:color'), 'auto')
                    shd_cbs.set(qn('w:fill'), 'FFFED5')
                    rPr_cbs.append(shd_cbs)
                    cb_close = para_cont.add_run(']')
                    cb_close.font.bold = True
                    cb_close.font.color.rgb = BRACKET_COLOR
                    rPr_cbc = cb_close._element.get_or_add_rPr()
                    shd_cbc = OxmlElement('w:shd')
                    shd_cbc.set(qn('w:val'), 'clear')
                    shd_cbc.set(qn('w:color'), 'auto')
                    shd_cbc.set(qn('w:fill'), 'FFFED5')
                    rPr_cbc.append(shd_cbc)
                    cont_end = para_cont.add_run(' and attach that sheet to this form]')
                    cont_end.font.name = 'Arial'
                    cont_end.font.size = Pt(12)
                    break

            # Hospital entry box with gold brackets
            hospital_text = self.hospital_treatment.text().strip()
            if not hospital_text and self.prac_address.text().strip():
                hospital_text = self.prac_address.text().strip()
            # Find hospital paragraph (usually around para 29)
            for i in range(25, min(40, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                prev_txt = ''.join(run.text for run in paragraphs[i-1].runs) if i > 0 else ''
                # Look for entry box after "following hospital"
                if "following hospital" in prev_txt or (not txt.strip() and i > 27):
                    set_entry_box(paragraphs[i], hospital_text if hospital_text else "")
                    break
            else:
                set_entry_box(paragraphs[29], hospital_text if hospital_text else "")

            # Signature line with gold brackets: Signed[ ] Date[ ]
            sig_date = self.sig_date.date().toString("dd MMMM yyyy")
            for i in range(30, min(45, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if "signed" in txt:
                    para_sig = paragraphs[i]
                    for run in para_sig.runs:
                        run.text = ""
                    while len(para_sig.runs) > 1:
                        para_sig._element.remove(para_sig.runs[-1]._element)
                    if para_sig.runs:
                        para_sig.runs[0].text = 'Signed'
                        para_sig.runs[0].font.name = 'Arial'
                        para_sig.runs[0].font.size = Pt(12)
                    else:
                        r = para_sig.add_run('Signed')
                        r.font.name = 'Arial'
                        r.font.size = Pt(12)
                    sig_open = para_sig.add_run('[')
                    sig_open.font.bold = True
                    sig_open.font.color.rgb = BRACKET_COLOR
                    rPr_so = sig_open._element.get_or_add_rPr()
                    shd_so = OxmlElement('w:shd')
                    shd_so.set(qn('w:val'), 'clear')
                    shd_so.set(qn('w:color'), 'auto')
                    shd_so.set(qn('w:fill'), 'FFFED5')
                    rPr_so.append(shd_so)
                    sig_content = para_sig.add_run('                                        ')
                    rPr_sc = sig_content._element.get_or_add_rPr()
                    shd_sc = OxmlElement('w:shd')
                    shd_sc.set(qn('w:val'), 'clear')
                    shd_sc.set(qn('w:color'), 'auto')
                    shd_sc.set(qn('w:fill'), 'FFFED5')
                    rPr_sc.append(shd_sc)
                    sig_close = para_sig.add_run(']')
                    sig_close.font.bold = True
                    sig_close.font.color.rgb = BRACKET_COLOR
                    date_label = para_sig.add_run(' Date')
                    date_label.font.name = 'Arial'
                    date_label.font.size = Pt(12)
                    date_open = para_sig.add_run('[')
                    date_open.font.bold = True
                    date_open.font.color.rgb = BRACKET_COLOR
                    rPr_do = date_open._element.get_or_add_rPr()
                    shd_do = OxmlElement('w:shd')
                    shd_do.set(qn('w:val'), 'clear')
                    shd_do.set(qn('w:color'), 'auto')
                    shd_do.set(qn('w:fill'), 'FFFED5')
                    rPr_do.append(shd_do)
                    date_content = para_sig.add_run(sig_date)
                    rPr_dc = date_content._element.get_or_add_rPr()
                    shd_dc = OxmlElement('w:shd')
                    shd_dc.set(qn('w:val'), 'clear')
                    shd_dc.set(qn('w:color'), 'auto')
                    shd_dc.set(qn('w:fill'), 'FFFED5')
                    rPr_dc.append(shd_dc)
                    date_close = para_sig.add_run(']')
                    date_close.font.bold = True
                    date_close.font.color.rgb = BRACKET_COLOR
                    break

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form A8 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        gender = "neutral"
        if self.gender_male.isChecked():
            gender = "male"
        elif self.gender_female.isChecked():
            gender = "female"

        return {
            "age": self.age_spin.value(),
            "gender": gender,
            "ethnicity": self.ethnicity_combo.currentText(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "prac_name": self.prac_name.text(),
            "prac_address": self.prac_address.text(),
            "prac_email": self.prac_email.text(),
            "prac_exam_date": self.prac_exam_date.date().toString("yyyy-MM-dd"),
            "prac_acquaintance": self.prac_acquaintance.isChecked(),
            "prac_section12": self.prac_section12.isChecked(),
            "dx_primary": self.dx_primary.currentText(),
            "dx_secondary": self.dx_secondary.currentText(),
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "self_hist_neglect": self.self_hist_neglect.isChecked(),
            "self_hist_risky": self.self_hist_risky.isChecked(),
            "self_hist_harm": self.self_hist_harm.isChecked(),
            "self_curr_neglect": self.self_curr_neglect.isChecked(),
            "self_curr_risky": self.self_curr_risky.isChecked(),
            "self_curr_harm": self.self_curr_harm.isChecked(),
            "others": self.others_cb.isChecked(),
            "others_hist_violence": self.others_hist_violence.isChecked(),
            "others_hist_verbal": self.others_hist_verbal.isChecked(),
            "others_hist_sexual": self.others_hist_sexual.isChecked(),
            "others_hist_stalking": self.others_hist_stalking.isChecked(),
            "others_hist_arson": self.others_hist_arson.isChecked(),
            "others_curr_violence": self.others_curr_violence.isChecked(),
            "others_curr_verbal": self.others_curr_verbal.isChecked(),
            "others_curr_sexual": self.others_curr_sexual.isChecked(),
            "others_curr_stalking": self.others_curr_stalking.isChecked(),
            "others_curr_arson": self.others_curr_arson.isChecked(),
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "clinical_reasons": self.clinical_reasons.text(),
            "hospital_treatment": self.hospital_treatment.text(),
            "sig_date": self.sig_date.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        if not state:
            return

        self.age_spin.setValue(state.get("age", 0))
        g = state.get("gender", "neutral")
        if g == "male":
            self.gender_male.setChecked(True)
        elif g == "female":
            self.gender_female.setChecked(True)
        else:
            self.gender_other.setChecked(True)

        idx = self.ethnicity_combo.findText(state.get("ethnicity", "Not specified"))
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)

        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.prac_name.setText(state.get("prac_name", ""))
        self.prac_address.setText(state.get("prac_address", ""))
        self.prac_email.setText(state.get("prac_email", ""))
        if state.get("prac_exam_date"):
            self.prac_exam_date.setDate(QDate.fromString(state["prac_exam_date"], "yyyy-MM-dd"))
        self.prac_acquaintance.setChecked(state.get("prac_acquaintance", False))
        self.prac_section12.setChecked(state.get("prac_section12", False))

        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)
        dx_secondary = state.get("dx_secondary", "")
        idx = self.dx_secondary.findText(dx_secondary)
        if idx >= 0:
            self.dx_secondary.setCurrentIndex(idx)
        else:
            self.dx_secondary.setCurrentText(dx_secondary)

        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.self_hist_neglect.setChecked(state.get("self_hist_neglect", False))
        self.self_hist_risky.setChecked(state.get("self_hist_risky", False))
        self.self_hist_harm.setChecked(state.get("self_hist_harm", False))
        self.self_curr_neglect.setChecked(state.get("self_curr_neglect", False))
        self.self_curr_risky.setChecked(state.get("self_curr_risky", False))
        self.self_curr_harm.setChecked(state.get("self_curr_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.others_hist_violence.setChecked(state.get("others_hist_violence", False))
        self.others_hist_verbal.setChecked(state.get("others_hist_verbal", False))
        self.others_hist_sexual.setChecked(state.get("others_hist_sexual", False))
        self.others_hist_stalking.setChecked(state.get("others_hist_stalking", False))
        self.others_hist_arson.setChecked(state.get("others_hist_arson", False))
        self.others_curr_violence.setChecked(state.get("others_curr_violence", False))
        self.others_curr_verbal.setChecked(state.get("others_curr_verbal", False))
        self.others_curr_sexual.setChecked(state.get("others_curr_sexual", False))
        self.others_curr_stalking.setChecked(state.get("others_curr_stalking", False))
        self.others_curr_arson.setChecked(state.get("others_curr_arson", False))
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        self.clinical_reasons.setText(state.get("clinical_reasons", ""))
        self.hospital_treatment.setText(state.get("hospital_treatment", ""))

        if state.get("sig_date"):
            self.sig_date.setDate(QDate.fromString(state["sig_date"], "yyyy-MM-dd"))

        # Update all cards
        self._update_patient_card()
        self._update_practitioner_card()
        self._update_clinical_card()
        self._update_signature_card()

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[A8Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[A8Form] Set gender: {gender}")
        # Set age if not already set (from age field or calculate from DOB)
        if hasattr(self, 'age_spin') and self.age_spin.value() == 0:
            age = patient_info.get("age")
            if not age and patient_info.get("dob"):
                from datetime import datetime
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if age and 0 < age < 120:
                self.age_spin.setValue(age)
                print(f"[A8Form] Set age: {age}")
        # Set ethnicity if not already selected
        if patient_info.get("ethnicity") and hasattr(self, 'ethnicity_combo'):
            current = self.ethnicity_combo.currentText()
            if current in ("Ethnicity", "Not specified", "Select..."):
                idx = self.ethnicity_combo.findText(patient_info["ethnicity"], Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    self.ethnicity_combo.setCurrentIndex(idx)
                    print(f"[A8Form] Set ethnicity: {patient_info['ethnicity']}")
