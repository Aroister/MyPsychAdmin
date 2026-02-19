#!/usr/bin/env python3
"""
HCR-20 V3 DOCX Exporter
Generates a Word document matching the exact format of the original HCR-20 template.
"""

from docx import Document
from docx.shared import Pt, Inches, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


def set_paragraph_border(paragraph, border_type="double", size=4, space=1):
    """Add double border to a paragraph (for the header box effect)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:val="{border_type}" w:sz="{size}" w:space="{space}" w:color="auto"/>'
        f'<w:left w:val="{border_type}" w:sz="{size}" w:space="4" w:color="auto"/>'
        f'<w:bottom w:val="{border_type}" w:sz="{size}" w:space="{space}" w:color="auto"/>'
        f'<w:right w:val="{border_type}" w:sz="{size}" w:space="25" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_tab_stop(paragraph, position_twips, alignment=WD_TAB_ALIGNMENT.LEFT):
    """Add a tab stop to a paragraph."""
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Twips(position_twips), alignment)


def create_bordered_field(doc, label, value, first=False):
    """Create a bordered paragraph with label and value using tab."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0

    # Add tab stop at ~8cm (4536 twips)
    add_tab_stop(para, 4536)

    # Set border
    set_paragraph_border(para)

    # Add label (bold)
    if label:
        run = para.add_run(f"{label}")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)
        run.bold = True
        para.add_run("\t")

    # Add value (not bold)
    if value:
        run = para.add_run(value)
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)
        run.bold = False

    return para


def create_bordered_empty(doc):
    """Create an empty bordered paragraph for spacing."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    set_paragraph_border(para)
    return para


def add_presence_relevance_table(doc, presence, relevance):
    """Add a presence/relevance summary table after an item."""
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'

    # Row 1: Presence
    cell = table.cell(0, 0)
    cell.text = "Presence"
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            run.bold = True

    cell = table.cell(0, 1)
    cell.text = presence or ""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

    # Row 2: Relevance
    cell = table.cell(1, 0)
    cell.text = "Relevance"
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            run.bold = True

    cell = table.cell(1, 1)
    cell.text = relevance or ""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

    doc.add_paragraph()
    return table


def add_section_text(doc, text, bold=False, underline=False, font_size=11):
    """Add a paragraph of text with formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)

    run = para.add_run(text)
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.underline = underline

    return para


def add_item_header(doc, item_code, item_title):
    """Add an HCR-20 item header (e.g., ITEM H1)."""
    # Item code
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(f"ITEM {item_code}")
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True

    # Item title (underlined)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(item_title)
    run.font.name = 'Arial Narrow'
    run.font.size = Pt(11)
    run.bold = True
    run.underline = True

    return para


def export_hcr20_docx(data: dict, output_path: str) -> bool:
    """
    Export HCR-20 data to a properly formatted Word document.

    Args:
        data: Dictionary containing all form data with keys like:
            - patient_name, dob, age, nhs_number, address
            - admission_date, legal_status
            - author_original, author_update, supervisor, review_to
            - date_original, date_update, date_next
            - sources
            - h1, h2, ... h10 (each with 'content', 'presence', 'relevance')
            - c1, c2, ... c5
            - r1, r2, ... r5
            - formulation, scenarios, management sections
            - signature_name, signature_role, signature_date
        output_path: Path to save the document

    Returns:
        True if successful, False otherwise
    """
    try:
        doc = Document()

        # Set default font for document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial Narrow'
        font.size = Pt(11)

        # ===== TITLE =====
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(12)
        run = title.add_run("HCR-20 Assessment Report")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(16)
        run.bold = True

        # ===== HEADER BOX =====
        # Empty bordered line at top
        create_bordered_empty(doc)

        # Patient details
        create_bordered_field(doc, "NAME:", data.get('patient_name', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "D.O.B:", data.get('dob', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "AGE:", data.get('age', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "NHS NUMBER:", data.get('nhs_number', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "ADDRESS:", data.get('address', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "DATE OF ADMISSION:", data.get('admission_date', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "LEGAL STATUS:", data.get('legal_status', ''))
        create_bordered_empty(doc)

        # Assessment details
        para = doc.add_paragraph()
        set_paragraph_border(para)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run("STRUCTURED CLINICAL JUDGMENT")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)
        run.bold = True

        create_bordered_field(doc, "TOOLS INCLUDED IN THIS REPORT:", "HCR-20 V3")
        create_bordered_empty(doc)
        create_bordered_field(doc, "AUTHOR OF ORIGINAL REPORT:", data.get('author_original', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "AUTHOR OF UPDATE REPORTS:", data.get('author_update', ''))

        # Supervisor line
        para = doc.add_paragraph()
        set_paragraph_border(para)
        para.paragraph_format.space_after = Pt(0)
        add_tab_stop(para, 4536)
        run = para.add_run("Under the Supervision of:")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)
        run.bold = True
        run.italic = True
        para.add_run("\t")
        run = para.add_run(data.get('supervisor', ''))
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)

        create_bordered_empty(doc)
        create_bordered_field(doc, "REPORT SENT FOR REVIEW TO:", data.get('review_to', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "Date of original report:", data.get('date_original', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "Date of update report:", data.get('date_update', ''))
        create_bordered_empty(doc)
        create_bordered_field(doc, "Date next update due:", data.get('date_next', ''))
        create_bordered_empty(doc)

        # Confidentiality notice
        doc.add_paragraph()
        para = doc.add_paragraph()
        run = para.add_run("This report is CONFIDENTIAL and should be restricted to persons with involvement with the patient. Sections of this report should not be cut and pasted into future reports without the expressed permission of the author.")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(10)
        run.italic = True

        # Page break
        doc.add_page_break()

        # ===== HCR-20 V3 INTRODUCTION =====
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("HCR-20 V3: RISK ASSESSMENT REPORT")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(14)
        run.bold = True

        doc.add_paragraph()

        # Introduction text
        para = doc.add_paragraph()
        run = para.add_run("HCR-20 V3 (Douglas et al., 2013)")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)
        run.bold = True

        para = doc.add_paragraph()
        run = para.add_run("The HCR-20 V3 is comprised of 20 risk factors across three subscales: ")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(11)
        run = para.add_run("Historical")
        run.font.name = 'Arial Narrow'
        run.bold = True
        para.add_run(", ")
        run = para.add_run("Clinical")
        run.font.name = 'Arial Narrow'
        run.bold = True
        para.add_run(", and ")
        run = para.add_run("Risk Management")
        run.font.name = 'Arial Narrow'
        run.bold = True
        para.add_run(".")

        para = doc.add_paragraph()
        para.add_run("The presence of factors is coded using a 3-level response format:").font.name = 'Arial Narrow'

        for rating in ["N = absent", "P = possibly or partially present", "Y = definitely present"]:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            run = para.add_run(f"• {rating}")
            run.font.name = 'Arial Narrow'

        para = doc.add_paragraph()
        run = para.add_run('"Omit" is used when there is no reliable information by which to judge the presence of the factor.')
        run.font.name = 'Arial Narrow'

        para = doc.add_paragraph()
        para.add_run('Evaluators then assess whether each risk factor is "relevant" to an individual\'s risk for violent behaviour.').font.name = 'Arial Narrow'

        para = doc.add_paragraph()
        run = para.add_run("The relevance of each factor is defined as '")
        run.font.name = 'Arial Narrow'
        run = para.add_run("Low")
        run.bold = True
        para.add_run("', '")
        run = para.add_run("Moderate")
        run.bold = True
        para.add_run("' or '")
        run = para.add_run("High")
        run.bold = True
        para.add_run("'.")

        doc.add_paragraph()

        # ===== SOURCES OF INFORMATION =====
        add_section_text(doc, "Sources of information:", bold=True)

        sources = data.get('sources', '')
        if sources:
            # Create a 2-column table for sources
            sources_list = [s.strip() for s in sources.split('\n') if s.strip()]
            if sources_list:
                rows = (len(sources_list) + 1) // 2
                table = doc.add_table(rows=max(1, rows), cols=2)
                table.style = 'Table Grid'

                for i, source in enumerate(sources_list):
                    row_idx = i // 2
                    col_idx = i % 2
                    if row_idx < len(table.rows):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = source.lstrip('•-● ').strip()
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = 'Arial Narrow'
                                run.font.size = Pt(11)

        doc.add_paragraph()

        # ===== HISTORICAL ITEMS =====
        section_title = doc.add_paragraph()
        run = section_title.add_run("Historical items")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(14)
        run.bold = True
        run.underline = True

        historical_items = [
            ("H1", "History of Problems with Violence"),
            ("H2", "History of Problems with Other Antisocial Behaviour"),
            ("H3", "History of Problems with Relationships"),
            ("H4", "History of Problems with Employment"),
            ("H5", "History of Problems with Substance Use"),
            ("H6", "History of Problems with Major Mental Disorder"),
            ("H7", "History of Problems with Personality Disorder"),
            ("H8", "History of Problems with Traumatic Experiences"),
            ("H9", "History of Problems with Violent Attitudes"),
            ("H10", "History of Problems with Treatment or Supervision Response"),
        ]

        for code, title in historical_items:
            key = code.lower()
            item_data = data.get(key, {})

            add_item_header(doc, code, title)

            # Content
            content = item_data.get('content', '') if isinstance(item_data, dict) else str(item_data)
            if content:
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)

            # Presence/Relevance table
            presence = item_data.get('presence', '') if isinstance(item_data, dict) else ''
            relevance = item_data.get('relevance', '') if isinstance(item_data, dict) else ''
            add_presence_relevance_table(doc, presence, relevance)

        # ===== CLINICAL ITEMS =====
        doc.add_page_break()
        section_title = doc.add_paragraph()
        run = section_title.add_run("Clinical items")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(14)
        run.bold = True
        run.underline = True

        clinical_items = [
            ("C1", "Recent Problems with Insight"),
            ("C2", "Recent Problems with Violent Ideation or Intent"),
            ("C3", "Recent Problems with Symptoms of Major Mental Disorder"),
            ("C4", "Recent Problems with Instability"),
            ("C5", "Recent Problems with Treatment or Supervision Response"),
        ]

        for code, title in clinical_items:
            key = code.lower()
            item_data = data.get(key, {})

            add_item_header(doc, code, title)

            content = item_data.get('content', '') if isinstance(item_data, dict) else str(item_data)
            if content:
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)

            presence = item_data.get('presence', '') if isinstance(item_data, dict) else ''
            relevance = item_data.get('relevance', '') if isinstance(item_data, dict) else ''
            add_presence_relevance_table(doc, presence, relevance)

        # ===== RISK MANAGEMENT ITEMS =====
        doc.add_page_break()
        section_title = doc.add_paragraph()
        run = section_title.add_run("Risk Management items")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(14)
        run.bold = True
        run.underline = True

        risk_items = [
            ("R1", "Future Problems with Professional Services and Plans"),
            ("R2", "Future Problems with Living Situation"),
            ("R3", "Future Problems with Personal Support"),
            ("R4", "Future Problems with Treatment or Supervision Response"),
            ("R5", "Future Problems with Stress or Coping"),
        ]

        for code, title in risk_items:
            key = code.lower()
            item_data = data.get(key, {})

            add_item_header(doc, code, title)

            content = item_data.get('content', '') if isinstance(item_data, dict) else str(item_data)
            if content:
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)

            presence = item_data.get('presence', '') if isinstance(item_data, dict) else ''
            relevance = item_data.get('relevance', '') if isinstance(item_data, dict) else ''
            add_presence_relevance_table(doc, presence, relevance)

        # ===== VIOLENCE RISK FORMULATION =====
        doc.add_page_break()
        section_title = doc.add_paragraph()
        run = section_title.add_run("Violence risk formulation")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(14)
        run.bold = True
        run.underline = True

        formulation = data.get('formulation', '')
        if formulation:
            para = doc.add_paragraph()
            run = para.add_run(formulation)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

        # Scenarios
        scenario_sections = [
            ("Scenarios (what kind of violence is likely to be committed, victims and likely motivation):", "scenario_nature"),
            ("Seriousness (psychological/physical harm to victims, could this escalate to a serious or life-threatening level?):", "scenario_severity"),
            ("Imminence (how soon could the individual engage in violence? What are the warning signs that risk is increasing or imminent?):", "scenario_imminence"),
            ("Frequency/Likelihood (how often might this violence occur? Is the risk chronic or acute?):", "scenario_frequency"),
        ]

        for label, key in scenario_sections:
            add_section_text(doc, label, bold=True)
            content = data.get(key, '')
            if content:
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)

        # ===== PROPOSED MANAGEMENT STRATEGIES =====
        doc.add_paragraph()
        section_title = doc.add_paragraph()
        run = section_title.add_run("Proposed management strategies")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(12)
        run.bold = True
        run.underline = True

        management_sections = [
            ("Risk-enhancing factors:", "risk_enhancing"),
            ("Protective factors:", "protective"),
            ("Monitoring:", "monitoring"),
        ]

        for label, key in management_sections:
            add_section_text(doc, label, bold=True)
            content = data.get(key, '')
            if content:
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.font.name = 'Arial Narrow'
                run.font.size = Pt(11)

        # ===== TREATMENT / RECOMMENDATIONS =====
        doc.add_paragraph()
        section_title = doc.add_paragraph()
        run = section_title.add_run("Treatment/ Recommendations")
        run.font.name = 'Arial Narrow'
        run.font.size = Pt(12)
        run.bold = True
        run.underline = True

        treatment = data.get('treatment', '')
        if treatment:
            para = doc.add_paragraph()
            run = para.add_run(treatment)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

        # Supervision
        add_section_text(doc, "Supervision", bold=True)
        supervision = data.get('supervision', '')
        if supervision:
            para = doc.add_paragraph()
            run = para.add_run(supervision)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

        # Victim safety planning
        add_section_text(doc, "Victim safety planning", bold=True)
        victim_safety = data.get('victim_safety', '')
        if victim_safety:
            para = doc.add_paragraph()
            run = para.add_run(victim_safety)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

        # ===== SIGNATURE =====
        doc.add_paragraph()
        doc.add_paragraph()

        sig_name = data.get('signature_name', '')
        if sig_name:
            para = doc.add_paragraph()
            run = para.add_run(sig_name)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            run.bold = True

        sig_role = data.get('signature_role', '')
        if sig_role:
            para = doc.add_paragraph()
            run = para.add_run(sig_role)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

        sig_date = data.get('signature_date', '')
        if sig_date:
            para = doc.add_paragraph()
            run = para.add_run(sig_date)
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)

        # Save document
        doc.save(output_path)
        return True

    except Exception as e:
        print(f"Error exporting HCR-20 document: {e}")
        import traceback
        traceback.print_exc()
        return False


if __name__ == "__main__":
    # Test with sample data
    test_data = {
        'patient_name': 'Test Patient',
        'dob': '1st January 1990',
        'age': '35',
        'nhs_number': '123 456 7890',
        'address': 'Test Unit, Test Hospital, Test Road, Test City, AB1 2CD',
        'admission_date': '1st January 2025',
        'legal_status': 'Section 3 of the Mental Health Act 1983 (amended 2007)',
        'author_original': 'Dr Test Author (Forensic Psychologist)',
        'author_update': 'Test Updater (Assistant Psychologist)',
        'supervisor': 'Dr Test Supervisor (Senior Clinical Psychologist)',
        'review_to': 'Test MDT',
        'date_original': 'January 2024',
        'date_update': 'January 2025',
        'date_next': 'July 2025',
        'sources': '• Previous HCR-20 report\n• Care notes\n• Tribunal hearing documents',
        'h1': {'content': 'Test content for H1...', 'presence': 'Present', 'relevance': 'High relevance'},
        'formulation': 'Test formulation content...',
        'signature_name': 'Dr Test Author',
        'signature_role': 'Forensic Psychologist',
        'signature_date': '24.01.25',
    }

    output = "/Users/avie/Desktop/HCR-20_Test_Output.docx"
    if export_hcr20_docx(test_data, output):
        print(f"Test document created: {output}")
    else:
        print("Failed to create test document")
