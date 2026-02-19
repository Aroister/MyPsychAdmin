# ================================================================
#  CTO1 FORM PAGE — Community Treatment Order (Card Layout)
#  Mental Health Act 1983 - Form CTO1 Regulation 6(1)(a), (b) and 6(2)(a)
#  Section 17A — Community treatment order
# ================================================================

from __future__ import annotations
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate, QTime
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit, QTimeEdit,
    QCheckBox, QPushButton, QSizePolicy, QFileDialog,
    QMessageBox, QToolButton, QRadioButton, QButtonGroup,
    QComboBox, QSpinBox, QCompleter, QStyleFactory, QSlider,
    QStackedWidget, QSplitter
)
from utils.resource_path import resource_path


# ================================================================
# RESIZABLE SECTION WITH DRAG BAR
# ================================================================
class ResizableSection(QFrame):
    """Section with a draggable bottom edge to resize height."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self._min_height = 60
        self._max_height = 400
        self._content = None
        self._dragging = False
        self._drag_start_y = 0
        self._drag_start_height = 0

        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(0, 0, 0, 0)
        self._layout.setSpacing(0)

        # Content container
        self._content_container = QWidget()
        self._content_layout = QVBoxLayout(self._content_container)
        self._content_layout.setContentsMargins(12, 8, 12, 8)
        self._layout.addWidget(self._content_container, 1)

        # Drag handle at bottom
        self._drag_handle = QFrame()
        self._drag_handle.setFixedHeight(8)
        self._drag_handle.setCursor(Qt.CursorShape.SizeVerCursor)
        self._drag_handle.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
                border-radius: 2px;
            }
            QFrame:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:0.3 #0891b2, stop:0.7 #0891b2, stop:1 transparent);
            }
        """)
        self._drag_handle.mousePressEvent = self._handle_press
        self._drag_handle.mouseMoveEvent = self._handle_move
        self._drag_handle.mouseReleaseEvent = self._handle_release
        self._layout.addWidget(self._drag_handle)

    def set_content(self, widget):
        self._content = widget
        self._content_layout.addWidget(widget)

    def set_content_height(self, h):
        h = max(self._min_height, min(self._max_height, h))
        self._content_container.setFixedHeight(h)

    def _handle_press(self, event):
        self._dragging = True
        self._drag_start_y = event.globalPosition().y()
        self._drag_start_height = self._content_container.height()

    def _handle_move(self, event):
        if self._dragging:
            delta = event.globalPosition().y() - self._drag_start_y
            new_height = self._drag_start_height + delta
            new_height = max(self._min_height, min(self._max_height, new_height))
            self._content_container.setFixedHeight(int(new_height))

    def _handle_release(self, event):
        self._dragging = False

# ICD-10 data
try:
    from icd10_dict import load_icd10_dict
    from icd10_curated import ICD10_GROUPED, ICD10_FLAT
except:
    ICD10_GROUPED = []
    ICD10_FLAT = []

from shared_widgets import create_zoom_row
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# NO-WHEEL SLIDER
# ================================================================
class NoWheelSlider(QSlider):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# NO-WHEEL COMBOBOX
# ================================================================
class NoWheelComboBox(QComboBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelSpinBox(QSpinBox):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelDateEdit(QDateEdit):
    def wheelEvent(self, event):
        event.ignore()

class NoWheelTimeEdit(QTimeEdit):
    def wheelEvent(self, event):
        event.ignore()


# ================================================================
# CTO1 CARD WIDGET
# ================================================================
class CTO1CardWidget(QFrame):
    """Clickable card with editable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("cto1Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#cto1Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#cto1Card:hover {
                border-color: #0891b2;
                background: #f0fdfa;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(12, 10, 12, 10)
        layout.setSpacing(6)

        # Header row with title and zoom controls
        header_row = QHBoxLayout()
        header_row.setContentsMargins(0, 0, 0, 0)
        header_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setStyleSheet("font-size: 19px; font-weight: 700; color: #0891b2;")
        header_row.addWidget(self.title_label)
        header_row.addStretch()

        layout.addLayout(header_row)

        self.content = MyPsychAdminRichTextEditor()
        self.content.setPlaceholderText("Click to edit...")
        self.content.setStyleSheet("""
            QTextEdit {
                font-size: 19px;
                color: #374151;
                background: transparent;
                border: none;
                padding: 0;
            }
        """)
        self.content.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        layout.addWidget(self.content, 1)

        # Add zoom controls to header row
        zoom_row = create_zoom_row(self.content, base_size=13)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                header_row.addWidget(item.widget())

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)

    def set_active(self, active: bool):
        """Set active state - shows persistent highlight."""
        if active:
            self.setStyleSheet("""
                QFrame#cto1Card {
                    background: #f0fdfa;
                    border: 2px solid #0891b2;
                    border-radius: 12px;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#cto1Card {
                    background: white;
                    border: 1px solid #e5e7eb;
                    border-radius: 12px;
                }
                QFrame#cto1Card:hover {
                    border-color: #0891b2;
                    background: #f0fdfa;
                }
                QLabel {
                    background: transparent;
                    border: none;
                }
            """)

    def set_content_text(self, text: str):
        self.content.setPlainText(text)

    def get_content(self) -> str:
        return self.content.toPlainText()


# ================================================================
# CTO1 TOOLBAR
# ================================================================
class CTO1Toolbar(QWidget):
    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            CTO1Toolbar { background: rgba(200, 215, 220, 0.95); border-bottom: 1px solid rgba(0,0,0,0.12); }
            QToolButton { background: transparent; color: #333; padding: 6px 10px; border-radius: 6px; font-size: 13px; font-weight: 500; }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; } QToolButton:hover { background: #1d4ed8; }")
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# CTO1 FORM PAGE
# ================================================================
class CTO1FormPage(QWidget):
    """CTO1 Form with card/popup layout."""

    go_back = Signal()

    ETHNICITIES = [
        "Afro-Caribbean", "Asian", "Caucasian", "Middle Eastern", "Mixed Race", "Not specified"
    ]

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self.cards = {}
        self._my_details = self._load_my_details()
        self._setup_ui()
        self._prefill()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {"full_name": details[1] or "", "email": details[7] or ""}

    def _prefill(self):
        if self._my_details.get("full_name"):
            self.rc_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.rc_email.setText(self._my_details["email"])

    def _get_pronouns(self):
        if self.gender_male.isChecked():
            return {"subj": "He", "subj_l": "he", "obj": "him", "pos": "His", "pos_l": "his"}
        elif self.gender_female.isChecked():
            return {"subj": "She", "subj_l": "she", "obj": "her", "pos": "Her", "pos_l": "her"}
        return {"subj": "They", "subj_l": "they", "obj": "them", "pos": "Their", "pos_l": "their"}

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(50)
        header.setStyleSheet("background: #0891b2;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 6px 14px;
                border-radius: 5px;
                font-size: 18px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self.go_back.emit)
        header_layout.addWidget(back_btn)

        title = QLabel("Form CTO1 — Community Treatment Order")
        title.setStyleSheet("font-size: 17px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Content area with splitter for resizable panels
        content = QWidget()
        content.setStyleSheet("background: #f9fafb;")
        content_layout = QVBoxLayout(content)
        content_layout.setContentsMargins(16, 16, 16, 16)
        content_layout.setSpacing(0)

        # Horizontal splitter for cards | popup
        self.main_splitter = QSplitter(Qt.Orientation.Horizontal)
        self.main_splitter.setHandleWidth(8)
        self.main_splitter.setStyleSheet("""
            QSplitter::handle {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.4 #d1d5db, stop:0.6 #d1d5db, stop:1 transparent);
                border-radius: 2px;
            }
            QSplitter::handle:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 transparent, stop:0.3 #0891b2, stop:0.7 #0891b2, stop:1 transparent);
            }
        """)

        # Left: Cards
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setMinimumWidth(200)
        cards_scroll.setMaximumWidth(550)
        cards_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: transparent;")
        self.cards_layout = QVBoxLayout(cards_container)
        self.cards_layout.setContentsMargins(16, 16, 16, 16)
        self.cards_layout.setSpacing(8)

        # Create cards
        self._create_details_card()
        self._create_grounds_card()
        self._create_conditions_card()
        self._create_amhp_card()
        self._create_signatures_card()

        self.cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        self.main_splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_stack = QStackedWidget()
        self.popup_stack.setMinimumWidth(400)
        self.popup_stack.setMaximumWidth(750)
        self.popup_stack.setStyleSheet("""
            QStackedWidget {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
        """)

        # Create popups
        self._create_details_popup()
        self._create_grounds_popup()
        self._create_conditions_popup()
        self._create_amhp_popup()
        self._create_signatures_popup()

        # Initialize cards with default date values
        self._update_amhp_card()
        self._update_signatures_card()

        self.main_splitter.addWidget(self.popup_stack)
        self.main_splitter.setSizes([280, 600])  # Initial sizes
        content_layout.addWidget(self.main_splitter)
        main_layout.addWidget(content, 1)

        # Connect live preview
        self._connect_grounds_live_preview()

        # Show first popup
        self._on_card_clicked("details")

    def _on_card_clicked(self, key: str):
        index_map = {"details": 0, "grounds": 1, "conditions": 2, "amhp": 3, "signatures": 4}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])
            for k, card in self.cards.items():
                card.set_active(k == key)

    # ----------------------------------------------------------------
    # CARDS
    # ----------------------------------------------------------------
    def _create_details_card(self):
        section = ResizableSection()
        section.set_content_height(170)
        section._min_height = 120
        section._max_height = 320

        card = CTO1CardWidget("Patient / RC Details", "details")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["details"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_grounds_card(self):
        section = ResizableSection()
        section.set_content_height(220)
        section._min_height = 140
        section._max_height = 450

        card = CTO1CardWidget("Grounds for CTO", "grounds")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["grounds"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_conditions_card(self):
        section = ResizableSection()
        section.set_content_height(170)
        section._min_height = 100
        section._max_height = 350

        card = CTO1CardWidget("Conditions", "conditions")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["conditions"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_amhp_card(self):
        section = ResizableSection()
        section.set_content_height(150)
        section._min_height = 100
        section._max_height = 280

        card = CTO1CardWidget("AMHP Agreement", "amhp")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["amhp"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    def _create_signatures_card(self):
        section = ResizableSection()
        section.set_content_height(150)
        section._min_height = 100
        section._max_height = 280

        card = CTO1CardWidget("Effective Date & Signatures", "signatures")
        self._hook_editor_focus(card.content)
        card.clicked.connect(self._on_card_clicked)
        self.cards["signatures"] = card
        section.set_content(card)
        self.cards_layout.addWidget(section)

    # ----------------------------------------------------------------
    # HELPER METHODS
    # ----------------------------------------------------------------
    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 18px;
            }
            QLineEdit:focus { border-color: #0891b2; }
        """)
        return edit

    def _create_date_edit(self) -> NoWheelDateEdit:
        date_edit = NoWheelDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 18px;
            }
        """)
        return date_edit

    def _create_time_edit(self) -> NoWheelTimeEdit:
        time_edit = NoWheelTimeEdit()
        time_edit.setTime(QTime.currentTime())
        time_edit.setStyleSheet("""
            QTimeEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 18px;
            }
        """)
        return time_edit

    def _create_styled_frame(self, color: str) -> QFrame:
        colors = {
            "cyan": ("#ecfeff", "#a5f3fc"),
            "green": ("#f0fdf4", "#bbf7d0"),
            "blue": ("#eff6ff", "#bfdbfe"),
            "red": ("#fef2f2", "#fecaca"),
            "purple": ("#faf5ff", "#e9d5ff"),
            "yellow": ("#fefce8", "#fef08a"),
        }
        bg, border = colors.get(color, ("#f9fafb", "#e5e7eb"))
        frame = QFrame()
        frame.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        frame.setStyleSheet(f"QFrame {{ background: {bg}; border: none; border-radius: 8px; }}")
        return frame


    # ----------------------------------------------------------------
    # POPUPS
    # ----------------------------------------------------------------
    def _create_details_popup(self):
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        header = QLabel("Patient / RC Details")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Patient section
        patient_frame = self._create_styled_frame("cyan")
        pf_layout = QVBoxLayout(patient_frame)
        pf_layout.setContentsMargins(12, 10, 12, 10)
        pf_layout.setSpacing(8)

        pf_header = QLabel("Patient")
        pf_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #0e7490;")
        pf_layout.addWidget(pf_header)

        self.patient_name = self._create_line_edit("Patient full name")
        pf_layout.addWidget(self.patient_name)

        self.patient_address = self._create_line_edit("Patient address")
        pf_layout.addWidget(self.patient_address)

        # Demographics row (age, gender, ethnicity - visible for clinical reasons)
        demo_row = QHBoxLayout()
        demo_row.setSpacing(8)

        # Age label and spin
        age_lbl = QLabel("Age:")
        age_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        demo_row.addWidget(age_lbl)
        self.age_spin = NoWheelSpinBox()
        self.age_spin.setRange(0, 120)
        self.age_spin.setFixedWidth(55)
        self.age_spin.setStyleSheet("QSpinBox { padding: 6px; border: 1px solid #d1d5db; border-radius: 4px; }")
        demo_row.addWidget(self.age_spin)

        self.gender_group = QButtonGroup(self)
        self.gender_male = QRadioButton("M")
        self.gender_female = QRadioButton("F")
        self.gender_other = QRadioButton("O")
        for rb in [self.gender_male, self.gender_female, self.gender_other]:
            rb.setStyleSheet("""
                QRadioButton {
                    font-size: 17px;
                    background: transparent;
                }
                QRadioButton::indicator {
                    width: 16px;
                    height: 16px;
                }
            """)
            self.gender_group.addButton(rb)
            demo_row.addWidget(rb)

        # Ethnicity combo - visible for clinical reasons
        self.ethnicity_combo = NoWheelComboBox()
        self.ethnicity_combo.addItem("Ethnicity")
        self.ethnicity_combo.addItems(self.ETHNICITIES)
        self.ethnicity_combo.setFixedWidth(130)
        self.ethnicity_combo.setStyleSheet("QComboBox { padding: 6px; font-size: 17px; border: 1px solid #d1d5db; border-radius: 4px; }")
        demo_row.addWidget(self.ethnicity_combo)
        demo_row.addStretch()

        pf_layout.addLayout(demo_row)
        layout.addWidget(patient_frame)

        # RC section
        rc_frame = self._create_styled_frame("green")
        rc_layout = QVBoxLayout(rc_frame)
        rc_layout.setContentsMargins(12, 10, 12, 10)
        rc_layout.setSpacing(8)

        rc_header = QLabel("Responsible Clinician")
        rc_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #166534;")
        rc_layout.addWidget(rc_header)

        self.rc_name = self._create_line_edit("RC full name")
        rc_layout.addWidget(self.rc_name)

        rc_row = QHBoxLayout()
        rc_row.setSpacing(8)
        self.rc_address = self._create_line_edit("RC address")
        rc_row.addWidget(self.rc_address, 2)
        self.rc_email = self._create_line_edit("Email")
        rc_row.addWidget(self.rc_email, 1)
        rc_layout.addLayout(rc_row)

        layout.addWidget(rc_frame)

        layout.addStretch()

        # Auto-sync to card
        self.patient_name.textChanged.connect(self._update_details_card)
        self.patient_address.textChanged.connect(self._update_details_card)
        self.rc_name.textChanged.connect(self._update_details_card)
        self.rc_address.textChanged.connect(self._update_details_card)

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_grounds_popup(self):
        popup = QWidget()
        popup.setStyleSheet("background: white;")
        main_layout = QVBoxLayout(popup)
        main_layout.setContentsMargins(20, 16, 20, 16)
        main_layout.setSpacing(12)

        header = QLabel("Grounds for CTO")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1f2937;")
        main_layout.addWidget(header)

        # Hidden label for state storage (preview removed - card auto-syncs)
        self.grounds_preview_label = QLabel("")
        self.grounds_preview_label.hide()

        # Scrollable form content
        form_scroll = QScrollArea()
        form_scroll.setWidgetResizable(True)
        form_scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        form_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        form_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        form_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        form_container = QWidget()
        form_container.setStyleSheet("background: transparent;")
        form_container.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        layout = QVBoxLayout(form_container)
        layout.setContentsMargins(0, 0, 8, 0)
        layout.setSpacing(10)

        # Mental Disorder
        md_frame = self._create_styled_frame("green")
        md_layout = QVBoxLayout(md_frame)
        md_layout.setContentsMargins(12, 10, 12, 10)
        md_layout.setSpacing(6)

        md_header = QLabel("Mental Disorder (ICD-10)")
        md_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #166534;")
        md_layout.addWidget(md_header)

        # Primary diagnosis dropdown with grouped items
        self.dx_primary = NoWheelComboBox()
        self.dx_primary.setEditable(True)
        self.dx_primary.lineEdit().setPlaceholderText("Select primary diagnosis...")
        self.dx_primary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_primary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_primary.addItem(f"── {group_name} ──", None)
            idx = self.dx_primary.count() - 1
            self.dx_primary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_primary.addItem(dx, dx)
        completer = QCompleter(ICD10_FLAT, self.dx_primary)
        completer.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_primary.setCompleter(completer)
        self.dx_primary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        md_layout.addWidget(self.dx_primary)

        # Secondary diagnosis dropdown
        self.dx_secondary = NoWheelComboBox()
        self.dx_secondary.setEditable(True)
        self.dx_secondary.lineEdit().setPlaceholderText("Secondary diagnosis (optional)...")
        self.dx_secondary.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.dx_secondary.addItem("", None)
        for group_name, diagnoses in ICD10_GROUPED:
            self.dx_secondary.addItem(f"── {group_name} ──", None)
            idx = self.dx_secondary.count() - 1
            self.dx_secondary.model().item(idx).setEnabled(False)
            for dx in diagnoses:
                self.dx_secondary.addItem(dx, dx)
        completer2 = QCompleter(ICD10_FLAT, self.dx_secondary)
        completer2.setCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        completer2.setFilterMode(Qt.MatchFlag.MatchContains)
        self.dx_secondary.setCompleter(completer2)
        self.dx_secondary.setStyleSheet("QComboBox { font-size: 17px; padding: 6px; border: 1px solid #d1d5db; border-radius: 6px; background: white; }")
        md_layout.addWidget(self.dx_secondary)

        layout.addWidget(md_frame)

        # Legal Criteria
        lc_frame = self._create_styled_frame("blue")
        lc_layout = QVBoxLayout(lc_frame)
        lc_layout.setContentsMargins(12, 10, 12, 10)
        lc_layout.setSpacing(4)

        lc_header = QLabel("Legal Criteria")
        lc_header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1e40af;")
        lc_layout.addWidget(lc_header)

        # Nature
        self.nature_cb = QCheckBox("Nature")
        self.nature_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.nature_cb.toggled.connect(self._on_nature_toggled)
        lc_layout.addWidget(self.nature_cb)

        self.nature_options = QWidget()
        nature_opt_layout = QVBoxLayout(self.nature_options)
        nature_opt_layout.setContentsMargins(16, 2, 0, 2)
        nature_opt_layout.setSpacing(2)

        self.relapsing_cb = QCheckBox("Relapsing and remitting")
        self.treatment_resistant_cb = QCheckBox("Treatment resistant")
        self.chronic_cb = QCheckBox("Chronic and enduring")
        for cb in [self.relapsing_cb, self.treatment_resistant_cb, self.chronic_cb]:
            cb.setStyleSheet("font-size: 17px; color: #6b7280;")
            nature_opt_layout.addWidget(cb)

        self.nature_options.hide()
        lc_layout.addWidget(self.nature_options)

        # Degree
        self.degree_cb = QCheckBox("Degree")
        self.degree_cb.setStyleSheet("font-size: 18px; font-weight: 600; color: #374151;")
        self.degree_cb.toggled.connect(self._on_degree_toggled)
        lc_layout.addWidget(self.degree_cb)

        self.degree_options = QWidget()
        degree_opt_layout = QVBoxLayout(self.degree_options)
        degree_opt_layout.setContentsMargins(16, 2, 0, 2)
        degree_opt_layout.setSpacing(4)

        slider_row = QHBoxLayout()
        self.degree_slider = NoWheelSlider(Qt.Orientation.Horizontal)
        self.degree_slider.setMinimum(1)
        self.degree_slider.setMaximum(4)
        self.degree_slider.setValue(2)
        self.degree_slider.setFixedWidth(80)
        slider_row.addWidget(self.degree_slider)
        self.degree_level_label = QLabel("Several")
        self.degree_level_label.setStyleSheet("font-size: 16px; color: #374151;")
        slider_row.addWidget(self.degree_level_label)
        slider_row.addStretch()
        degree_opt_layout.addLayout(slider_row)

        self.degree_details = QLineEdit()
        self.degree_details.setPlaceholderText("Symptoms including...")
        self.degree_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        degree_opt_layout.addWidget(self.degree_details)

        self.degree_options.hide()
        lc_layout.addWidget(self.degree_options)

        # Necessity
        nec_lbl = QLabel("Necessity:")
        nec_lbl.setStyleSheet("font-size: 17px; font-weight: 600; color: #374151; margin-top: 4px;")
        lc_layout.addWidget(nec_lbl)

        # Health
        self.health_cb = QCheckBox("Health")
        self.health_cb.setStyleSheet("font-size: 17px; color: #374151;")
        self.health_cb.toggled.connect(self._on_health_toggled)
        lc_layout.addWidget(self.health_cb)

        self.health_options = QWidget()
        health_opt_layout = QVBoxLayout(self.health_options)
        health_opt_layout.setContentsMargins(16, 2, 0, 2)
        health_opt_layout.setSpacing(2)

        self.mental_health_cb = QCheckBox("Mental Health")
        self.mental_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.mental_health_cb.toggled.connect(self._on_mental_health_toggled)
        health_opt_layout.addWidget(self.mental_health_cb)

        self.mh_options = QWidget()
        mh_opt_layout = QVBoxLayout(self.mh_options)
        mh_opt_layout.setContentsMargins(16, 2, 0, 2)
        mh_opt_layout.setSpacing(2)
        self.poor_compliance_cb = QCheckBox("Poor compliance")
        self.limited_insight_cb = QCheckBox("Limited insight")
        for cb in [self.poor_compliance_cb, self.limited_insight_cb]:
            cb.setStyleSheet("font-size: 16px; color: #9ca3af;")
            mh_opt_layout.addWidget(cb)
        self.mh_options.hide()
        health_opt_layout.addWidget(self.mh_options)

        self.physical_health_cb = QCheckBox("Physical Health")
        self.physical_health_cb.setStyleSheet("font-size: 16px; color: #6b7280;")
        self.physical_health_cb.toggled.connect(self._on_physical_health_toggled)
        health_opt_layout.addWidget(self.physical_health_cb)

        self.physical_health_details = QLineEdit()
        self.physical_health_details.setPlaceholderText("Physical health details...")
        self.physical_health_details.setStyleSheet("font-size: 16px; padding: 4px; border: 1px solid #d1d5db; border-radius: 4px;")
        self.physical_health_details.hide()
        health_opt_layout.addWidget(self.physical_health_details)

        self.health_options.hide()
        lc_layout.addWidget(self.health_options)

        # Safety
        self.safety_cb = QCheckBox("Safety")
        self.safety_cb.setStyleSheet("font-size: 17px; color: #374151;")
        self.safety_cb.toggled.connect(self._on_safety_toggled)
        lc_layout.addWidget(self.safety_cb)

        self.safety_options = QWidget()
        safety_opt_layout = QVBoxLayout(self.safety_options)
        safety_opt_layout.setContentsMargins(16, 2, 0, 2)
        safety_opt_layout.setSpacing(4)

        # Self
        self.self_harm_cb = QCheckBox("Self")
        self.self_harm_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.self_harm_cb.toggled.connect(self._on_self_toggled)
        safety_opt_layout.addWidget(self.self_harm_cb)

        self.self_options = QWidget()
        self_opt_layout = QVBoxLayout(self.self_options)
        self_opt_layout.setContentsMargins(16, 2, 0, 2)
        self_opt_layout.setSpacing(2)

        self_hist_lbl = QLabel("Historical:")
        self_hist_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_hist_lbl)
        self.self_hist_neglect = QCheckBox("Self neglect")
        self.self_hist_risky = QCheckBox("Risky situations")
        self.self_hist_harm = QCheckBox("Self harm")
        for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm]:
            cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
            self_opt_layout.addWidget(cb)

        self_curr_lbl = QLabel("Current:")
        self_curr_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600;")
        self_opt_layout.addWidget(self_curr_lbl)
        self.self_curr_neglect = QCheckBox("Self neglect")
        self.self_curr_risky = QCheckBox("Risky situations")
        self.self_curr_harm = QCheckBox("Self harm")
        for cb in [self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
            cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
            self_opt_layout.addWidget(cb)

        self.self_options.hide()
        safety_opt_layout.addWidget(self.self_options)

        # Others
        self.others_cb = QCheckBox("Others")
        self.others_cb.setStyleSheet("font-size: 16px; font-weight: 600; color: #6b7280;")
        self.others_cb.toggled.connect(self._on_others_toggled)
        safety_opt_layout.addWidget(self.others_cb)

        self.others_options = QWidget()
        others_opt_layout = QVBoxLayout(self.others_options)
        others_opt_layout.setContentsMargins(16, 2, 0, 2)
        others_opt_layout.setSpacing(2)

        others_hist_lbl = QLabel("Historical:")
        others_hist_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_hist_lbl)
        self.others_hist_violence = QCheckBox("Violence")
        self.others_hist_verbal = QCheckBox("Verbal aggression")
        self.others_hist_sexual = QCheckBox("Sexual violence")
        self.others_hist_stalking = QCheckBox("Stalking")
        self.others_hist_arson = QCheckBox("Arson")
        for cb in [self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual, self.others_hist_stalking, self.others_hist_arson]:
            cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
            others_opt_layout.addWidget(cb)

        others_curr_lbl = QLabel("Current:")
        others_curr_lbl.setStyleSheet("font-size: 11px; color: #9ca3af; font-weight: 600;")
        others_opt_layout.addWidget(others_curr_lbl)
        self.others_curr_violence = QCheckBox("Violence")
        self.others_curr_verbal = QCheckBox("Verbal aggression")
        self.others_curr_sexual = QCheckBox("Sexual violence")
        self.others_curr_stalking = QCheckBox("Stalking")
        self.others_curr_arson = QCheckBox("Arson")
        for cb in [self.others_curr_violence, self.others_curr_verbal, self.others_curr_sexual, self.others_curr_stalking, self.others_curr_arson]:
            cb.setStyleSheet("font-size: 11px; color: #9ca3af;")
            others_opt_layout.addWidget(cb)

        self.others_options.hide()
        safety_opt_layout.addWidget(self.others_options)

        self.safety_options.hide()
        lc_layout.addWidget(self.safety_options)

        layout.addWidget(lc_frame)

        # CTO Appropriate
        cto_frame = self._create_styled_frame("red")
        cto_layout = QVBoxLayout(cto_frame)
        cto_layout.setContentsMargins(12, 10, 12, 10)
        cto_layout.setSpacing(4)

        cto_header = QLabel("CTO Appropriate Because")
        cto_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #991b1b;")
        cto_layout.addWidget(cto_header)

        self.tried_failed_cb = QCheckBox("Tried/Failed discharge")
        self.insight_cb = QCheckBox("Insight into illness")
        self.compliance_cb = QCheckBox("Compliance with treatment")
        self.supervision_cb = QCheckBox("Needs community supervision")
        for cb in [self.tried_failed_cb, self.insight_cb, self.compliance_cb, self.supervision_cb]:
            cb.setStyleSheet("font-size: 16px; color: #374151;")
            cto_layout.addWidget(cb)

        layout.addWidget(cto_frame)

        # RC Signature date
        sig_frame = self._create_styled_frame("purple")
        sig_layout = QHBoxLayout(sig_frame)
        sig_layout.setContentsMargins(12, 10, 12, 10)
        sig_lbl = QLabel("RC Signature Date:")
        sig_lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        sig_layout.addWidget(sig_lbl)
        self.rc_sig_date = self._create_date_edit()
        self.rc_sig_date.setFixedWidth(160)
        sig_layout.addWidget(self.rc_sig_date)
        sig_layout.addStretch()

        layout.addWidget(sig_frame)
        layout.addStretch()

        form_scroll.setWidget(form_container)
        main_layout.addWidget(form_scroll, 1)

        self.popup_stack.addWidget(popup)

    def _create_conditions_popup(self):
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        header = QLabel("Conditions (section 17B(2))")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Standard conditions checkboxes
        cond_frame = self._create_styled_frame("green")
        cond_layout = QVBoxLayout(cond_frame)
        cond_layout.setContentsMargins(12, 10, 12, 10)
        cond_layout.setSpacing(6)

        cond_header = QLabel("Standard Conditions")
        cond_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #166534;")
        cond_layout.addWidget(cond_header)

        self.cond_cmht_cb = QCheckBox("See CMHT")
        self.cond_medication_cb = QCheckBox("Comply with medication")
        self.cond_residence_cb = QCheckBox("Residence")
        for cb in [self.cond_cmht_cb, self.cond_medication_cb, self.cond_residence_cb]:
            cb.setStyleSheet("font-size: 17px; color: #374151;")
            cb.toggled.connect(self._update_conditions_preview)
            cond_layout.addWidget(cb)

        layout.addWidget(cond_frame)

        # Hidden label for state storage (preview removed - card auto-syncs)
        self.conditions_preview = QLabel("")
        self.conditions_preview.hide()

        layout.addStretch()

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_amhp_popup(self):
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        header = QLabel("AMHP Agreement")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        amhp_frame = self._create_styled_frame("purple")
        af_layout = QVBoxLayout(amhp_frame)
        af_layout.setContentsMargins(12, 10, 12, 10)
        af_layout.setSpacing(8)

        self.amhp_name = self._create_line_edit("AMHP full name")
        af_layout.addWidget(self.amhp_name)

        self.amhp_address = self._create_line_edit("AMHP address")
        af_layout.addWidget(self.amhp_address)

        self.amhp_authority = self._create_line_edit("Acting on behalf of (Local Social Services Authority)")
        af_layout.addWidget(self.amhp_authority)

        self.amhp_approved_by = self._create_line_edit("Approved by (if different)")
        af_layout.addWidget(self.amhp_approved_by)

        sig_row = QHBoxLayout()
        sig_lbl = QLabel("Signature Date:")
        sig_lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.amhp_sig_date = self._create_date_edit()
        self.amhp_sig_date.setFixedWidth(160)
        sig_row.addWidget(self.amhp_sig_date)
        sig_row.addStretch()
        af_layout.addLayout(sig_row)

        layout.addWidget(amhp_frame)

        layout.addStretch()

        # Auto-sync to card
        self.amhp_name.textChanged.connect(self._update_amhp_card)
        self.amhp_address.textChanged.connect(self._update_amhp_card)
        self.amhp_authority.textChanged.connect(self._update_amhp_card)
        self.amhp_sig_date.dateChanged.connect(self._update_amhp_card)

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    def _create_signatures_popup(self):
        popup = QScrollArea()
        popup.setWidgetResizable(True)
        popup.setFrameShape(QScrollArea.Shape.NoFrame)
        popup.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)

        header = QLabel("Effective Date & Signatures")
        header.setStyleSheet("font-size: 18px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        eff_frame = self._create_styled_frame("red")
        ef_layout = QVBoxLayout(eff_frame)
        ef_layout.setContentsMargins(12, 10, 12, 10)
        ef_layout.setSpacing(8)

        ef_header = QLabel("CTO Effective From")
        ef_header.setStyleSheet("font-size: 17px; font-weight: 700; color: #991b1b;")
        ef_layout.addWidget(ef_header)

        row = QHBoxLayout()
        row.setSpacing(12)

        date_lbl = QLabel("Date:")
        date_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        row.addWidget(date_lbl)
        self.effective_date = self._create_date_edit()
        self.effective_date.setFixedWidth(160)
        row.addWidget(self.effective_date)

        time_lbl = QLabel("Time:")
        time_lbl.setStyleSheet("font-size: 17px; color: #374151;")
        row.addWidget(time_lbl)
        self.effective_time = self._create_time_edit()
        self.effective_time.setFixedWidth(120)
        row.addWidget(self.effective_time)
        row.addStretch()

        ef_layout.addLayout(row)

        sig_row = QHBoxLayout()
        sig_lbl = QLabel("RC Signature Date:")
        sig_lbl.setStyleSheet("font-size: 17px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.final_sig_date = self._create_date_edit()
        self.final_sig_date.setFixedWidth(160)
        sig_row.addWidget(self.final_sig_date)
        sig_row.addStretch()
        ef_layout.addLayout(sig_row)

        layout.addWidget(eff_frame)

        warning = QLabel("This order is NOT VALID unless all parts are completed and signed.")
        warning.setWordWrap(True)
        warning.setStyleSheet("font-size: 17px; color: #dc2626; font-weight: 600; padding: 8px; background: #fef2f2; border-radius: 6px;")
        layout.addWidget(warning)

        layout.addStretch()

        # Auto-sync to card
        self.effective_date.dateChanged.connect(self._update_signatures_card)
        self.effective_time.timeChanged.connect(self._update_signatures_card)
        self.final_sig_date.dateChanged.connect(self._update_signatures_card)

        popup.setWidget(container)
        self.popup_stack.addWidget(popup)

    # ----------------------------------------------------------------
    # TOGGLE HANDLERS
    # ----------------------------------------------------------------
    def _on_nature_toggled(self, checked):
        self.nature_options.setVisible(checked)
        if not checked:
            self.relapsing_cb.setChecked(False)
            self.treatment_resistant_cb.setChecked(False)
            self.chronic_cb.setChecked(False)

    def _on_degree_toggled(self, checked):
        self.degree_options.setVisible(checked)
        if not checked:
            self.degree_details.clear()

    def _on_health_toggled(self, checked):
        self.health_options.setVisible(checked)
        if not checked:
            self.mental_health_cb.setChecked(False)
            self.physical_health_cb.setChecked(False)

    def _on_mental_health_toggled(self, checked):
        self.mh_options.setVisible(checked)
        if not checked:
            self.poor_compliance_cb.setChecked(False)
            self.limited_insight_cb.setChecked(False)

    def _on_physical_health_toggled(self, checked):
        self.physical_health_details.setVisible(checked)
        if not checked:
            self.physical_health_details.clear()

    def _on_safety_toggled(self, checked):
        self.safety_options.setVisible(checked)
        if not checked:
            self.self_harm_cb.setChecked(False)
            self.others_cb.setChecked(False)

    def _on_self_toggled(self, checked):
        self.self_options.setVisible(checked)
        if not checked:
            for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm,
                       self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm]:
                cb.setChecked(False)

    def _on_others_toggled(self, checked):
        self.others_options.setVisible(checked)
        if not checked:
            for cb in [self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual,
                       self.others_hist_stalking, self.others_hist_arson,
                       self.others_curr_violence, self.others_curr_verbal, self.others_curr_sexual,
                       self.others_curr_stalking, self.others_curr_arson]:
                cb.setChecked(False)

    # ----------------------------------------------------------------
    # LIVE PREVIEW CONNECTION
    # ----------------------------------------------------------------
    def _connect_grounds_live_preview(self):
        self.dx_primary.currentTextChanged.connect(self._update_grounds_preview)
        self.dx_secondary.currentTextChanged.connect(self._update_grounds_preview)

        self.nature_cb.toggled.connect(self._update_grounds_preview)
        self.relapsing_cb.toggled.connect(self._update_grounds_preview)
        self.treatment_resistant_cb.toggled.connect(self._update_grounds_preview)
        self.chronic_cb.toggled.connect(self._update_grounds_preview)

        self.degree_cb.toggled.connect(self._update_grounds_preview)
        self.degree_slider.valueChanged.connect(self._update_grounds_preview)
        self.degree_details.textChanged.connect(self._update_grounds_preview)

        self.health_cb.toggled.connect(self._update_grounds_preview)
        self.mental_health_cb.toggled.connect(self._update_grounds_preview)
        self.poor_compliance_cb.toggled.connect(self._update_grounds_preview)
        self.limited_insight_cb.toggled.connect(self._update_grounds_preview)
        self.physical_health_cb.toggled.connect(self._update_grounds_preview)
        self.physical_health_details.textChanged.connect(self._update_grounds_preview)

        self.safety_cb.toggled.connect(self._update_grounds_preview)
        self.self_harm_cb.toggled.connect(self._update_grounds_preview)
        self.others_cb.toggled.connect(self._update_grounds_preview)

        for cb in [self.self_hist_neglect, self.self_hist_risky, self.self_hist_harm,
                   self.self_curr_neglect, self.self_curr_risky, self.self_curr_harm,
                   self.others_hist_violence, self.others_hist_verbal, self.others_hist_sexual,
                   self.others_hist_stalking, self.others_hist_arson,
                   self.others_curr_violence, self.others_curr_verbal, self.others_curr_sexual,
                   self.others_curr_stalking, self.others_curr_arson]:
            cb.toggled.connect(self._update_grounds_preview)

        self.tried_failed_cb.toggled.connect(self._update_grounds_preview)
        self.insight_cb.toggled.connect(self._update_grounds_preview)
        self.compliance_cb.toggled.connect(self._update_grounds_preview)
        self.supervision_cb.toggled.connect(self._update_grounds_preview)

        self.patient_name.textChanged.connect(self._update_grounds_preview)
        self.gender_male.toggled.connect(self._update_grounds_preview)
        self.gender_female.toggled.connect(self._update_grounds_preview)

        # Slider label update
        self.degree_slider.valueChanged.connect(self._on_degree_slider_changed)

    def _on_degree_slider_changed(self, value):
        levels = {1: "Some", 2: "Several", 3: "Many", 4: "Overwhelming"}
        self.degree_level_label.setText(levels.get(value, "Several"))

    def _update_grounds_preview(self):
        text = self._generate_grounds_text()
        self.grounds_preview_label.setText(text if text else "")
        self._update_grounds_card()

    def _update_conditions_preview(self):
        text = self._generate_conditions_text()
        self.conditions_preview.setText(text if text else "")
        self._update_conditions_card()

    # ----------------------------------------------------------------
    # TEXT GENERATION
    # ----------------------------------------------------------------
    def _generate_grounds_text(self) -> str:
        p = self._get_pronouns()
        patient_name = self.patient_name.text().strip()
        name_display = patient_name if patient_name else "The patient"

        paragraphs = []

        # Para 1: Demographics + Diagnosis
        para1_parts = []
        opening_parts = []
        if self.age_spin.value() > 0:
            opening_parts.append(f"{self.age_spin.value()} year old")
        ethnicity = self.ethnicity_combo.currentText()
        if ethnicity not in ("Ethnicity", "Not specified"):
            opening_parts.append(ethnicity.replace(" British", ""))
        if self.gender_male.isChecked():
            opening_parts.append("man")
        elif self.gender_female.isChecked():
            opening_parts.append("woman")

        diagnoses = []
        if self.dx_primary.currentText().strip():
            diagnoses.append(self.dx_primary.currentText().strip())
        if self.dx_secondary.currentText().strip():
            diagnoses.append(self.dx_secondary.currentText().strip())

        if diagnoses:
            demo_str = " ".join(opening_parts) if opening_parts else ""
            if demo_str:
                para1_parts.append(f"{name_display} is a {demo_str} who suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")
            else:
                para1_parts.append(f"{name_display} suffers from {diagnoses[0]} which is a mental disorder as defined by the Mental Health Act.")

            if self.nature_cb.isChecked() and self.degree_cb.isChecked():
                para1_parts.append("The disorder is both of a nature and degree which makes it appropriate for the patient to receive medical treatment.")
            elif self.nature_cb.isChecked():
                para1_parts.append("The disorder is of a nature which makes it appropriate for the patient to receive medical treatment.")
            elif self.degree_cb.isChecked():
                para1_parts.append("The disorder is of a degree which makes it appropriate for the patient to receive medical treatment.")

            if self.nature_cb.isChecked():
                nature_types = []
                if self.relapsing_cb.isChecked():
                    nature_types.append("relapsing and remitting")
                if self.treatment_resistant_cb.isChecked():
                    nature_types.append("treatment resistant")
                if self.chronic_cb.isChecked():
                    nature_types.append("chronic and enduring")
                if nature_types:
                    para1_parts.append(f"The nature of the illness is {' and '.join(nature_types)}.")

            if self.degree_cb.isChecked():
                levels = {1: "some", 2: "several", 3: "many", 4: "overwhelming"}
                level = levels.get(self.degree_slider.value(), "several")
                details = self.degree_details.text().strip()
                dx_name = diagnoses[0].split(" (")[0].lower() if diagnoses else "the disorder"
                if details:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms including {details}.")
                else:
                    para1_parts.append(f"The degree is evidenced by the presence of {level} symptoms.")

        if para1_parts:
            paragraphs.append(" ".join(para1_parts))

        # Para 2: Necessity
        para2_parts = []
        necessity_items = []
        if self.health_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} health")
        if self.safety_cb.isChecked() and self.self_harm_cb.isChecked():
            necessity_items.append(f"{p['pos_l']} own safety")
        if self.safety_cb.isChecked() and self.others_cb.isChecked():
            necessity_items.append("protection of others")

        if necessity_items:
            if len(necessity_items) == 1:
                para2_parts.append(f"A CTO is necessary for {necessity_items[0]}.")
            elif len(necessity_items) == 2:
                para2_parts.append(f"A CTO is necessary for {necessity_items[0]} and {necessity_items[1]}.")
            else:
                para2_parts.append(f"A CTO is necessary for {necessity_items[0]}, {necessity_items[1]} and {necessity_items[2]}.")

        # Health details
        if self.health_cb.isChecked() and self.mental_health_cb.isChecked():
            mh_items = []
            if self.poor_compliance_cb.isChecked():
                mh_items.append("poor compliance with treatment")
            if self.limited_insight_cb.isChecked():
                mh_items.append(f"limited insight into {p['pos_l']} illness")
            if mh_items:
                para2_parts.append(f"Regarding {p['pos_l']} health, I would be concerned about {' and '.join(mh_items)} resulting in a deterioration in mental state if not on a CTO.")

        if self.health_cb.isChecked() and self.physical_health_cb.isChecked():
            details = self.physical_health_details.text().strip()
            if details:
                para2_parts.append(f"I am also concerned about {p['pos_l']} physical health: {details}.")

        if para2_parts:
            paragraphs.append(" ".join(para2_parts))

        # Para 3: CTO Appropriateness
        cto_reasons = []
        if self.tried_failed_cb.isChecked():
            cto_reasons.append("previous attempts at discharge have failed without a CTO")
        if self.insight_cb.isChecked():
            cto_reasons.append(f"{p['subj_l']} has limited insight necessitating close monitoring under the mental health act")
        if self.compliance_cb.isChecked():
            cto_reasons.append(f"{p['subj']} needs community monitoring afforded by the CTO to ensure compliance with medication")
        if self.supervision_cb.isChecked():
            cto_reasons.append(f"{p['subj_l']} requires community supervision to maintain stability")

        if cto_reasons:
            if len(cto_reasons) == 1:
                paragraphs.append(f"A CTO is appropriate because {cto_reasons[0]}.")
            else:
                paragraphs.append(f"A CTO is appropriate because {cto_reasons[0]} and {cto_reasons[1]}.")

        return "\n\n".join(paragraphs)

    def _generate_conditions_text(self) -> str:
        conditions = []
        num = 1

        if self.cond_cmht_cb.isChecked():
            conditions.append(f"{num}. To comply with reviews as defined by the care-coordinator and the RC.")
            num += 1

        if self.cond_medication_cb.isChecked():
            conditions.append(f"{num}. To adhere to psychiatric medications as prescribed by the RC.")
            num += 1

        if self.cond_residence_cb.isChecked():
            conditions.append(f"{num}. To reside at an address in accordance with the requirements of the CMHT/RC.")
            num += 1

        return "\n".join(conditions)

    # ----------------------------------------------------------------
    # CARD UPDATE METHODS
    # ----------------------------------------------------------------
    def _update_details_card(self):
        parts = []
        if self.patient_name.text():
            parts.append(self.patient_name.text())
        if self.patient_address.text():
            parts.append(self.patient_address.text()[:40] + "...")
        if self.rc_name.text():
            parts.append(f"RC: {self.rc_name.text()}")
        if self.rc_address.text():
            parts.append(self.rc_address.text()[:40] + "...")
        self.cards["details"].set_content_text("\n".join(parts) if parts else "")

    def _update_grounds_card(self):
        text = self._generate_grounds_text()
        self.cards["grounds"].set_content_text(text if text else "")

    def _update_conditions_card(self):
        text = self._generate_conditions_text()
        self.cards["conditions"].set_content_text(text if text else "")

    def _update_amhp_card(self):
        parts = []
        if self.amhp_name.text():
            parts.append(self.amhp_name.text())
        if self.amhp_address.text():
            parts.append(self.amhp_address.text())
        if self.amhp_authority.text():
            parts.append(self.amhp_authority.text())
        parts.append(self.amhp_sig_date.date().toString('dd MMM yyyy'))
        self.cards["amhp"].set_content_text("\n".join(parts))

    def _update_signatures_card(self):
        parts = []
        parts.append(self.effective_date.date().toString('dd MMM yyyy'))
        parts.append(self.effective_time.time().toString('HH:mm'))
        parts.append(self.final_sig_date.date().toString('dd MMM yyyy'))
        self.cards["signatures"].set_content_text("\n".join(parts))

    # ----------------------------------------------------------------
    # EDITOR FOCUS TRACKING
    # ----------------------------------------------------------------
    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    # ----------------------------------------------------------------
    # FORM ACTIONS
    # ----------------------------------------------------------------
    def _clear_form(self):
        reply = QMessageBox.question(self, "Clear Form", "Clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.patient_name.clear()
            self.patient_address.clear()
            self.age_spin.setValue(0)
            self.gender_group.setExclusive(False)
            self.gender_male.setChecked(False)
            self.gender_female.setChecked(False)
            self.gender_other.setChecked(False)
            self.gender_group.setExclusive(True)
            self.ethnicity_combo.setCurrentIndex(0)
            self.rc_name.clear()
            self.rc_address.clear()
            self.rc_email.clear()
            self.dx_primary.setCurrentIndex(0)
            self.dx_secondary.setCurrentIndex(0)
            self.nature_cb.setChecked(False)
            self.degree_cb.setChecked(False)
            self.health_cb.setChecked(False)
            self.safety_cb.setChecked(False)
            self.tried_failed_cb.setChecked(False)
            self.insight_cb.setChecked(False)
            self.compliance_cb.setChecked(False)
            self.supervision_cb.setChecked(False)
            self.cond_cmht_cb.setChecked(False)
            self.cond_medication_cb.setChecked(False)
            self.cond_residence_cb.setChecked(False)
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_authority.clear()
            self.amhp_approved_by.clear()
            self.rc_sig_date.setDate(QDate.currentDate())
            self.amhp_sig_date.setDate(QDate.currentDate())
            self.effective_date.setDate(QDate.currentDate())
            self.effective_time.setTime(QTime.currentTime())
            self.final_sig_date.setDate(QDate.currentDate())
            for card in self.cards.values():
                card.set_content_text("")
            self._prefill()

    def _export_docx(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form CTO1",
            f"Form_CTO1_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_CTO1_template.docx')
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form CTO1 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Gold bracket color
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)

            # Clean ALL paragraphs - remove permission markers and convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), 'FFFED5')
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), 'FFFED5')

            def set_entry_box(para, content=""):
                """Set entry box with bold gold brackets [content]"""
                if not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr = bracket_open._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'FFFED5')
                rPr2.append(shd2)
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr3 = bracket_close._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)

            def highlight_yellow(para):
                """Rebuild paragraph with cream highlight"""
                text = ''.join(run.text for run in para.runs)
                if not text:
                    text = para.text
                if not text.strip():
                    return
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                if para.runs:
                    para.runs[0].text = text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                    rPr = para.runs[0]._element.get_or_add_rPr()
                    for old_shd in rPr.findall(qn('w:shd')):
                        rPr.remove(old_shd)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)
                else:
                    run = para.add_run(text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    rPr = run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            paragraphs = doc.paragraphs

            # Debug: Print paragraph indices
            print("DEBUG CTO1: Paragraphs 0-70:")
            for i in range(min(70, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)[:60]
                print(f"DEBUG: Para {i}: {txt}")

            # RC details entry box with gold brackets (para 5)
            rc_parts = []
            if self.rc_name.text().strip():
                rc_parts.append(self.rc_name.text().strip())
            if self.rc_address.text().strip():
                rc_parts.append(self.rc_address.text().strip())
            if self.rc_email.text().strip():
                rc_parts.append(self.rc_email.text().strip())
            set_entry_box(paragraphs[5], ", ".join(rc_parts) if rc_parts else "")

            # Patient details entry box with gold brackets (para 8)
            patient_text = self.patient_name.text().strip()
            if self.patient_address.text().strip():
                patient_text += ", " + self.patient_address.text().strip()
            set_entry_box(paragraphs[8], patient_text if patient_text else "")

            # Necessity options (i), (ii), (iii) - paras 12, 13, 14
            # Format: [the patient's health
            #         the patient's safety
            #         the protection of other persons]
            # Para 12: gold [ at start, then highlighted text
            p12 = paragraphs[12]
            p12_text = ''.join(run.text for run in p12.runs)
            for run in p12.runs:
                run.text = ""
            while len(p12.runs) > 1:
                p12._element.remove(p12.runs[-1]._element)
            if p12.runs:
                p12.runs[0].text = '['
                p12.runs[0].font.bold = True
                p12.runs[0].font.color.rgb = BRACKET_COLOR
            else:
                p12_open = p12.add_run('[')
                p12_open.font.bold = True
                p12_open.font.color.rgb = BRACKET_COLOR
            p12_content = p12.add_run(p12_text.strip())
            p12_content.font.name = 'Arial'
            p12_content.font.size = Pt(12)
            rPr12c = p12_content._element.get_or_add_rPr()
            shd12c = OxmlElement('w:shd')
            shd12c.set(qn('w:val'), 'clear')
            shd12c.set(qn('w:color'), 'auto')
            shd12c.set(qn('w:fill'), 'FFFED5')
            rPr12c.append(shd12c)

            # Para 13: just highlighted text (no brackets)
            p13 = paragraphs[13]
            p13_text = ''.join(run.text for run in p13.runs)
            for run in p13.runs:
                run.text = ""
            while len(p13.runs) > 1:
                p13._element.remove(p13.runs[-1]._element)
            if p13.runs:
                p13.runs[0].text = p13_text.strip()
                p13.runs[0].font.name = 'Arial'
                p13.runs[0].font.size = Pt(12)
                rPr13 = p13.runs[0]._element.get_or_add_rPr()
                shd13 = OxmlElement('w:shd')
                shd13.set(qn('w:val'), 'clear')
                shd13.set(qn('w:color'), 'auto')
                shd13.set(qn('w:fill'), 'FFFED5')
                rPr13.append(shd13)

            # Para 14: highlighted text, then gold ] at end
            p14 = paragraphs[14]
            p14_text = ''.join(run.text for run in p14.runs)
            for run in p14.runs:
                run.text = ""
            while len(p14.runs) > 1:
                p14._element.remove(p14.runs[-1]._element)
            if p14.runs:
                p14.runs[0].text = p14_text.strip()
                p14.runs[0].font.name = 'Arial'
                p14.runs[0].font.size = Pt(12)
                rPr14 = p14.runs[0]._element.get_or_add_rPr()
                shd14 = OxmlElement('w:shd')
                shd14.set(qn('w:val'), 'clear')
                shd14.set(qn('w:color'), 'auto')
                shd14.set(qn('w:fill'), 'FFFED5')
                rPr14.append(shd14)
            p14_close = p14.add_run(']')
            p14_close.font.bold = True
            p14_close.font.color.rgb = BRACKET_COLOR

            # Strikethrough unselected necessity options
            if not self.health_cb.isChecked():
                strikethrough_para(paragraphs[12])
            if not (self.safety_cb.isChecked() and self.self_harm_cb.isChecked()):
                strikethrough_para(paragraphs[13])
            if not (self.safety_cb.isChecked() and self.others_cb.isChecked()):
                strikethrough_para(paragraphs[14])

            # Grounds entry box with gold brackets
            grounds_text = self.cards["grounds"].get_content()
            set_entry_box(paragraphs[21], grounds_text if grounds_text.strip() else "")

            # Find and format "[If you need to continue...]" paragraphs
            for i in range(15, min(60, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "If you need to continue" in txt:
                    para_cont = paragraphs[i]
                    for run in para_cont.runs:
                        run.text = ""
                    while len(para_cont.runs) > 1:
                        para_cont._element.remove(para_cont.runs[-1]._element)
                    if para_cont.runs:
                        para_cont.runs[0].text = '[If you need to continue on a separate sheet please indicate here '
                        para_cont.runs[0].font.name = 'Arial'
                        para_cont.runs[0].font.size = Pt(12)
                    cb_open = para_cont.add_run('[')
                    cb_open.font.bold = True
                    cb_open.font.color.rgb = BRACKET_COLOR
                    rPr_cbo = cb_open._element.get_or_add_rPr()
                    shd_cbo = OxmlElement('w:shd')
                    shd_cbo.set(qn('w:val'), 'clear')
                    shd_cbo.set(qn('w:color'), 'auto')
                    shd_cbo.set(qn('w:fill'), 'FFFED5')
                    rPr_cbo.append(shd_cbo)
                    cb_space = para_cont.add_run(' ')
                    rPr_cbs = cb_space._element.get_or_add_rPr()
                    shd_cbs = OxmlElement('w:shd')
                    shd_cbs.set(qn('w:val'), 'clear')
                    shd_cbs.set(qn('w:color'), 'auto')
                    shd_cbs.set(qn('w:fill'), 'FFFED5')
                    rPr_cbs.append(shd_cbs)
                    cb_close = para_cont.add_run(']')
                    cb_close.font.bold = True
                    cb_close.font.color.rgb = BRACKET_COLOR
                    rPr_cbc = cb_close._element.get_or_add_rPr()
                    shd_cbc = OxmlElement('w:shd')
                    shd_cbc.set(qn('w:val'), 'clear')
                    shd_cbc.set(qn('w:color'), 'auto')
                    shd_cbc.set(qn('w:fill'), 'FFFED5')
                    rPr_cbc.append(shd_cbc)
                    cont_end = para_cont.add_run(' and attach that sheet to this form]')
                    cont_end.font.name = 'Arial'
                    cont_end.font.size = Pt(12)

            # Conditions entry box with gold brackets
            conditions_text = self.cards["conditions"].get_content()
            set_entry_box(paragraphs[30], conditions_text if conditions_text.strip() else "")

            # Find and format "[that authority]" paragraph
            for i in range(40, min(55, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "that authority" in txt.lower():
                    para_ta = paragraphs[i]
                    for run in para_ta.runs:
                        run.text = ""
                    while len(para_ta.runs) > 1:
                        para_ta._element.remove(para_ta.runs[-1]._element)
                    if para_ta.runs:
                        para_ta.runs[0].text = '['
                        para_ta.runs[0].font.bold = True
                        para_ta.runs[0].font.color.rgb = BRACKET_COLOR
                    ta_content = para_ta.add_run('that authority')
                    ta_content.font.name = 'Arial'
                    ta_content.font.size = Pt(12)
                    rPr_ta = ta_content._element.get_or_add_rPr()
                    shd_ta = OxmlElement('w:shd')
                    shd_ta.set(qn('w:val'), 'clear')
                    shd_ta.set(qn('w:color'), 'auto')
                    shd_ta.set(qn('w:fill'), 'FFFED5')
                    rPr_ta.append(shd_ta)
                    ta_close = para_ta.add_run(']')
                    ta_close.font.bold = True
                    ta_close.font.color.rgb = BRACKET_COLOR
                    break

            # RC signature line with gold brackets
            sig_date = self.rc_sig_date.date().toString("dd MMMM yyyy")
            for i in range(35, min(45, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if "signed" in txt and i < 45:
                    para_sig = paragraphs[i]
                    for run in para_sig.runs:
                        run.text = ""
                    while len(para_sig.runs) > 1:
                        para_sig._element.remove(para_sig.runs[-1]._element)
                    if para_sig.runs:
                        para_sig.runs[0].text = 'Signed'
                        para_sig.runs[0].font.name = 'Arial'
                        para_sig.runs[0].font.size = Pt(12)
                    sig_open = para_sig.add_run('[')
                    sig_open.font.bold = True
                    sig_open.font.color.rgb = BRACKET_COLOR
                    rPr_so = sig_open._element.get_or_add_rPr()
                    shd_so = OxmlElement('w:shd')
                    shd_so.set(qn('w:val'), 'clear')
                    shd_so.set(qn('w:color'), 'auto')
                    shd_so.set(qn('w:fill'), 'FFFED5')
                    rPr_so.append(shd_so)
                    sig_content = para_sig.add_run('                                        ')
                    rPr_sc = sig_content._element.get_or_add_rPr()
                    shd_sc = OxmlElement('w:shd')
                    shd_sc.set(qn('w:val'), 'clear')
                    shd_sc.set(qn('w:color'), 'auto')
                    shd_sc.set(qn('w:fill'), 'FFFED5')
                    rPr_sc.append(shd_sc)
                    sig_close = para_sig.add_run(']')
                    sig_close.font.bold = True
                    sig_close.font.color.rgb = BRACKET_COLOR
                    date_label = para_sig.add_run(' Date')
                    date_label.font.name = 'Arial'
                    date_label.font.size = Pt(12)
                    date_open = para_sig.add_run('[')
                    date_open.font.bold = True
                    date_open.font.color.rgb = BRACKET_COLOR
                    rPr_do = date_open._element.get_or_add_rPr()
                    shd_do = OxmlElement('w:shd')
                    shd_do.set(qn('w:val'), 'clear')
                    shd_do.set(qn('w:color'), 'auto')
                    shd_do.set(qn('w:fill'), 'FFFED5')
                    rPr_do.append(shd_do)
                    date_content = para_sig.add_run(sig_date)
                    rPr_dc = date_content._element.get_or_add_rPr()
                    shd_dc = OxmlElement('w:shd')
                    shd_dc.set(qn('w:val'), 'clear')
                    shd_dc.set(qn('w:color'), 'auto')
                    shd_dc.set(qn('w:fill'), 'FFFED5')
                    rPr_dc.append(shd_dc)
                    date_close = para_sig.add_run(']')
                    date_close.font.bold = True
                    date_close.font.color.rgb = BRACKET_COLOR
                    break

            # AMHP details entry box
            amhp_text = self.amhp_name.text().strip()
            if self.amhp_address.text().strip():
                amhp_text += ", " + self.amhp_address.text().strip()
            set_entry_box(paragraphs[41], amhp_text if amhp_text else "")

            # AMHP authority entry box
            if self.amhp_authority.text().strip():
                set_entry_box(paragraphs[43], self.amhp_authority.text().strip())
            else:
                set_entry_box(paragraphs[43], "")

            # AMHP approved by entry box
            if self.amhp_approved_by.text().strip():
                set_entry_box(paragraphs[47], self.amhp_approved_by.text().strip())
            else:
                set_entry_box(paragraphs[47], "")

            # PART 2 - AMHP signature line: Signed:[ ] Approved mental health professional
            amhp_sig_date = self.amhp_sig_date.date().toString("dd MMMM yyyy")
            for i in range(48, min(58, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if "signed" in txt:
                    para_amhp_sig = paragraphs[i]
                    for run in para_amhp_sig.runs:
                        run.text = ""
                    while len(para_amhp_sig.runs) > 1:
                        para_amhp_sig._element.remove(para_amhp_sig.runs[-1]._element)
                    if para_amhp_sig.runs:
                        para_amhp_sig.runs[0].text = 'Signed:'
                        para_amhp_sig.runs[0].font.name = 'Arial'
                        para_amhp_sig.runs[0].font.size = Pt(12)
                    asig_open = para_amhp_sig.add_run('[')
                    asig_open.font.bold = True
                    asig_open.font.color.rgb = BRACKET_COLOR
                    rPr_aso = asig_open._element.get_or_add_rPr()
                    shd_aso = OxmlElement('w:shd')
                    shd_aso.set(qn('w:val'), 'clear')
                    shd_aso.set(qn('w:color'), 'auto')
                    shd_aso.set(qn('w:fill'), 'FFFED5')
                    rPr_aso.append(shd_aso)
                    asig_content = para_amhp_sig.add_run('                                        ')
                    rPr_asc = asig_content._element.get_or_add_rPr()
                    shd_asc = OxmlElement('w:shd')
                    shd_asc.set(qn('w:val'), 'clear')
                    shd_asc.set(qn('w:color'), 'auto')
                    shd_asc.set(qn('w:fill'), 'FFFED5')
                    rPr_asc.append(shd_asc)
                    asig_close = para_amhp_sig.add_run(']')
                    asig_close.font.bold = True
                    asig_close.font.color.rgb = BRACKET_COLOR
                    asig_label = para_amhp_sig.add_run(' Approved mental health professional')
                    asig_label.font.name = 'Arial'
                    asig_label.font.size = Pt(12)
                    break

            # PART 2 - AMHP date line: Date[ ]
            for i in range(48, min(58, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if txt.strip().startswith("date"):
                    para_amhp_date = paragraphs[i]
                    for run in para_amhp_date.runs:
                        run.text = ""
                    while len(para_amhp_date.runs) > 1:
                        para_amhp_date._element.remove(para_amhp_date.runs[-1]._element)
                    if para_amhp_date.runs:
                        para_amhp_date.runs[0].text = 'Date'
                        para_amhp_date.runs[0].font.name = 'Arial'
                        para_amhp_date.runs[0].font.size = Pt(12)
                    adate_open = para_amhp_date.add_run('[')
                    adate_open.font.bold = True
                    adate_open.font.color.rgb = BRACKET_COLOR
                    rPr_ado = adate_open._element.get_or_add_rPr()
                    shd_ado = OxmlElement('w:shd')
                    shd_ado.set(qn('w:val'), 'clear')
                    shd_ado.set(qn('w:color'), 'auto')
                    shd_ado.set(qn('w:fill'), 'FFFED5')
                    rPr_ado.append(shd_ado)
                    adate_content = para_amhp_date.add_run(amhp_sig_date)
                    rPr_adc = adate_content._element.get_or_add_rPr()
                    shd_adc = OxmlElement('w:shd')
                    shd_adc.set(qn('w:val'), 'clear')
                    shd_adc.set(qn('w:color'), 'auto')
                    shd_adc.set(qn('w:fill'), 'FFFED5')
                    rPr_adc.append(shd_adc)
                    adate_close = para_amhp_date.add_run(']')
                    adate_close.font.bold = True
                    adate_close.font.color.rgb = BRACKET_COLOR
                    break

            # Effective date entry box
            eff_date = self.effective_date.date().toString("dd MMMM yyyy")
            set_entry_box(paragraphs[57], eff_date)

            # Effective time entry box
            eff_time = self.effective_time.time().toString("HH:mm")
            set_entry_box(paragraphs[59], eff_time)

            # Final signature line with gold brackets
            final_date = self.final_sig_date.date().toString("dd MMMM yyyy")
            for i in range(58, min(70, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if "signed" in txt:
                    para_fsig = paragraphs[i]
                    for run in para_fsig.runs:
                        run.text = ""
                    while len(para_fsig.runs) > 1:
                        para_fsig._element.remove(para_fsig.runs[-1]._element)
                    if para_fsig.runs:
                        para_fsig.runs[0].text = 'Signed:'
                        para_fsig.runs[0].font.name = 'Arial'
                        para_fsig.runs[0].font.size = Pt(12)
                    fsig_open = para_fsig.add_run('[')
                    fsig_open.font.bold = True
                    fsig_open.font.color.rgb = BRACKET_COLOR
                    rPr_fso = fsig_open._element.get_or_add_rPr()
                    shd_fso = OxmlElement('w:shd')
                    shd_fso.set(qn('w:val'), 'clear')
                    shd_fso.set(qn('w:color'), 'auto')
                    shd_fso.set(qn('w:fill'), 'FFFED5')
                    rPr_fso.append(shd_fso)
                    fsig_content = para_fsig.add_run('                                        ')
                    rPr_fsc = fsig_content._element.get_or_add_rPr()
                    shd_fsc = OxmlElement('w:shd')
                    shd_fsc.set(qn('w:val'), 'clear')
                    shd_fsc.set(qn('w:color'), 'auto')
                    shd_fsc.set(qn('w:fill'), 'FFFED5')
                    rPr_fsc.append(shd_fsc)
                    fsig_close = para_fsig.add_run(']')
                    fsig_close.font.bold = True
                    fsig_close.font.color.rgb = BRACKET_COLOR
                    fsig_label = para_fsig.add_run(' Responsible clinician')
                    fsig_label.font.name = 'Arial'
                    fsig_label.font.size = Pt(12)
                    break

            # Final date line
            for i in range(58, min(70, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs).lower()
                if txt.strip().startswith("date"):
                    para_fdate = paragraphs[i]
                    for run in para_fdate.runs:
                        run.text = ""
                    while len(para_fdate.runs) > 1:
                        para_fdate._element.remove(para_fdate.runs[-1]._element)
                    if para_fdate.runs:
                        para_fdate.runs[0].text = 'Date:'
                        para_fdate.runs[0].font.name = 'Arial'
                        para_fdate.runs[0].font.size = Pt(12)
                    fdate_open = para_fdate.add_run('[')
                    fdate_open.font.bold = True
                    fdate_open.font.color.rgb = BRACKET_COLOR
                    rPr_fdo = fdate_open._element.get_or_add_rPr()
                    shd_fdo = OxmlElement('w:shd')
                    shd_fdo.set(qn('w:val'), 'clear')
                    shd_fdo.set(qn('w:color'), 'auto')
                    shd_fdo.set(qn('w:fill'), 'FFFED5')
                    rPr_fdo.append(shd_fdo)
                    fdate_content = para_fdate.add_run(final_date)
                    rPr_fdc = fdate_content._element.get_or_add_rPr()
                    shd_fdc = OxmlElement('w:shd')
                    shd_fdc.set(qn('w:val'), 'clear')
                    shd_fdc.set(qn('w:color'), 'auto')
                    shd_fdc.set(qn('w:fill'), 'FFFED5')
                    rPr_fdc.append(shd_fdc)
                    fdate_close = para_fdate.add_run(']')
                    fdate_close.font.bold = True
                    fdate_close.font.color.rgb = BRACKET_COLOR
                    break

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form CTO1 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    # ----------------------------------------------------------------
    # STATE MANAGEMENT
    # ----------------------------------------------------------------
    def get_state(self) -> dict:
        return {
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "age": self.age_spin.value(),
            "gender": "male" if self.gender_male.isChecked() else "female" if self.gender_female.isChecked() else "other" if self.gender_other.isChecked() else "",
            "ethnicity": self.ethnicity_combo.currentText(),
            "rc_name": self.rc_name.text(),
            "rc_address": self.rc_address.text(),
            "rc_email": self.rc_email.text(),
            "dx_primary": self.dx_primary.currentText(),
            "dx_secondary": self.dx_secondary.currentText(),
            "nature": self.nature_cb.isChecked(),
            "relapsing": self.relapsing_cb.isChecked(),
            "treatment_resistant": self.treatment_resistant_cb.isChecked(),
            "chronic": self.chronic_cb.isChecked(),
            "degree": self.degree_cb.isChecked(),
            "degree_level": self.degree_slider.value(),
            "degree_details": self.degree_details.text(),
            "health": self.health_cb.isChecked(),
            "mental_health": self.mental_health_cb.isChecked(),
            "poor_compliance": self.poor_compliance_cb.isChecked(),
            "limited_insight": self.limited_insight_cb.isChecked(),
            "physical_health": self.physical_health_cb.isChecked(),
            "physical_health_details": self.physical_health_details.text(),
            "safety": self.safety_cb.isChecked(),
            "self_harm": self.self_harm_cb.isChecked(),
            "others": self.others_cb.isChecked(),
            "tried_failed": self.tried_failed_cb.isChecked(),
            "insight": self.insight_cb.isChecked(),
            "compliance": self.compliance_cb.isChecked(),
            "supervision": self.supervision_cb.isChecked(),
            "cond_cmht": self.cond_cmht_cb.isChecked(),
            "cond_medication": self.cond_medication_cb.isChecked(),
            "cond_residence": self.cond_residence_cb.isChecked(),
            "rc_sig_date": self.rc_sig_date.date().toString("yyyy-MM-dd"),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.text(),
            "amhp_authority": self.amhp_authority.text(),
            "amhp_approved_by": self.amhp_approved_by.text(),
            "amhp_sig_date": self.amhp_sig_date.date().toString("yyyy-MM-dd"),
            "effective_date": self.effective_date.date().toString("yyyy-MM-dd"),
            "effective_time": self.effective_time.time().toString("HH:mm"),
            "final_sig_date": self.final_sig_date.date().toString("yyyy-MM-dd"),
            "card_details": self.cards["details"].get_content(),
            "card_grounds": self.cards["grounds"].get_content(),
            "card_conditions": self.cards["conditions"].get_content(),
            "card_amhp": self.cards["amhp"].get_content(),
            "card_signatures": self.cards["signatures"].get_content(),
        }

    def set_state(self, state: dict):
        if not state:
            return
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.age_spin.setValue(state.get("age", 0))
        gender = state.get("gender", "")
        if gender == "male":
            self.gender_male.setChecked(True)
        elif gender == "female":
            self.gender_female.setChecked(True)
        elif gender == "other":
            self.gender_other.setChecked(True)
        idx = self.ethnicity_combo.findText(state.get("ethnicity", "Ethnicity"))
        if idx >= 0:
            self.ethnicity_combo.setCurrentIndex(idx)
        self.rc_name.setText(state.get("rc_name", ""))
        self.rc_address.setText(state.get("rc_address", ""))
        self.rc_email.setText(state.get("rc_email", ""))
        dx_primary = state.get("dx_primary", "")
        idx = self.dx_primary.findText(dx_primary)
        if idx >= 0:
            self.dx_primary.setCurrentIndex(idx)
        else:
            self.dx_primary.setCurrentText(dx_primary)
        dx_secondary = state.get("dx_secondary", "")
        idx = self.dx_secondary.findText(dx_secondary)
        if idx >= 0:
            self.dx_secondary.setCurrentIndex(idx)
        else:
            self.dx_secondary.setCurrentText(dx_secondary)
        self.nature_cb.setChecked(state.get("nature", False))
        self.relapsing_cb.setChecked(state.get("relapsing", False))
        self.treatment_resistant_cb.setChecked(state.get("treatment_resistant", False))
        self.chronic_cb.setChecked(state.get("chronic", False))
        self.degree_cb.setChecked(state.get("degree", False))
        self.degree_slider.setValue(state.get("degree_level", 2))
        self.degree_details.setText(state.get("degree_details", ""))
        self.health_cb.setChecked(state.get("health", False))
        self.mental_health_cb.setChecked(state.get("mental_health", False))
        self.poor_compliance_cb.setChecked(state.get("poor_compliance", False))
        self.limited_insight_cb.setChecked(state.get("limited_insight", False))
        self.physical_health_cb.setChecked(state.get("physical_health", False))
        self.physical_health_details.setText(state.get("physical_health_details", ""))
        self.safety_cb.setChecked(state.get("safety", False))
        self.self_harm_cb.setChecked(state.get("self_harm", False))
        self.others_cb.setChecked(state.get("others", False))
        self.tried_failed_cb.setChecked(state.get("tried_failed", False))
        self.insight_cb.setChecked(state.get("insight", False))
        self.compliance_cb.setChecked(state.get("compliance", False))
        self.supervision_cb.setChecked(state.get("supervision", False))
        self.cond_cmht_cb.setChecked(state.get("cond_cmht", False))
        self.cond_medication_cb.setChecked(state.get("cond_medication", False))
        self.cond_residence_cb.setChecked(state.get("cond_residence", False))
        if state.get("rc_sig_date"):
            self.rc_sig_date.setDate(QDate.fromString(state["rc_sig_date"], "yyyy-MM-dd"))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setText(state.get("amhp_address", ""))
        self.amhp_authority.setText(state.get("amhp_authority", ""))
        self.amhp_approved_by.setText(state.get("amhp_approved_by", ""))
        if state.get("amhp_sig_date"):
            self.amhp_sig_date.setDate(QDate.fromString(state["amhp_sig_date"], "yyyy-MM-dd"))
        if state.get("effective_date"):
            self.effective_date.setDate(QDate.fromString(state["effective_date"], "yyyy-MM-dd"))
        if state.get("effective_time"):
            self.effective_time.setTime(QTime.fromString(state["effective_time"], "HH:mm"))
        if state.get("final_sig_date"):
            self.final_sig_date.setDate(QDate.fromString(state["final_sig_date"], "yyyy-MM-dd"))

        # Restore card contents
        if state.get("card_details"):
            self.cards["details"].set_content_text(state["card_details"])
        if state.get("card_grounds"):
            self.cards["grounds"].set_content_text(state["card_grounds"])
        if state.get("card_conditions"):
            self.cards["conditions"].set_content_text(state["card_conditions"])
        if state.get("card_amhp"):
            self.cards["amhp"].set_content_text(state["card_amhp"])
        if state.get("card_signatures"):
            self.cards["signatures"].set_content_text(state["card_signatures"])

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[CTO1Form] Set patient name: {patient_info['name']}")
        # Set gender if not already selected
        if patient_info.get("gender"):
            gender = patient_info["gender"].lower()
            if not self.gender_male.isChecked() and not self.gender_female.isChecked() and not self.gender_other.isChecked():
                if gender == "male":
                    self.gender_male.setChecked(True)
                elif gender == "female":
                    self.gender_female.setChecked(True)
                print(f"[CTO1Form] Set gender: {gender}")
        # Set age if not already set (from age field or calculate from DOB)
        if hasattr(self, 'age_spin') and self.age_spin.value() == 0:
            age = patient_info.get("age")
            if not age and patient_info.get("dob"):
                from datetime import datetime
                dob = patient_info["dob"]
                if isinstance(dob, datetime):
                    today = datetime.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
            if age and 0 < age < 120:
                self.age_spin.setValue(age)
                print(f"[CTO1Form] Set age: {age}")
        # Set ethnicity if not already selected
        if patient_info.get("ethnicity") and hasattr(self, 'ethnicity_combo'):
            current = self.ethnicity_combo.currentText()
            if current in ("Ethnicity", "Not specified", "Select..."):
                idx = self.ethnicity_combo.findText(patient_info["ethnicity"], Qt.MatchFlag.MatchContains)
                if idx >= 0:
                    self.ethnicity_combo.setCurrentIndex(idx)
                    print(f"[CTO1Form] Set ethnicity: {patient_info['ethnicity']}")
