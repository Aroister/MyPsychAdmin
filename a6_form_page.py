# ================================================================
#  A6 FORM PAGE — Section 3 Application by AMHP
#  Mental Health Act 1983 - Form A6 Regulation 4(1)(c)(ii)
#  Card/Popup Layout with ResizableSection
# ================================================================

from __future__ import annotations
import re
from datetime import datetime
from PySide6.QtCore import Qt, Signal, QDate
from PySide6.QtGui import QColor
from PySide6.QtWidgets import (
    QGraphicsDropShadowEffect,
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QFrame,
    QScrollArea, QLineEdit, QTextEdit, QDateEdit,
    QRadioButton, QButtonGroup, QCheckBox, QPushButton,
    QSizePolicy, QFileDialog, QMessageBox, QGroupBox,
    QToolButton, QStackedWidget, QSplitter
)

from background_history_popup import ResizableSection
from shared_widgets import create_zoom_row
from utils.resource_path import resource_path
from mha_form_toolbar import MHAFormToolbar
from mypsy_richtext_editor import MyPsychAdminRichTextEditor


# ================================================================
# A6 CARD WIDGET
# ================================================================
class A6CardWidget(QFrame):
    """Card widget with fixed header and scrollable content."""
    clicked = Signal(str)

    def __init__(self, title: str, key: str, parent=None):
        super().__init__(parent)
        self.key = key
        self.setObjectName("a6Card")
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setStyleSheet("""
            QFrame#a6Card {
                background: white;
                border: 1px solid #e5e7eb;
                border-radius: 12px;
            }
            QFrame#a6Card:hover {
                border-color: #2563eb;
                background: #eff6ff;
            }
            QLabel {
                background: transparent;
                border: none;
            }
            QLabel#cardTitle {
                font-size: 20px;
                font-weight: 600;
                color: #2563eb;
            }
            QLabel#cardContent {
                font-size: 20px;
                color: #374151;
                padding: 4px;
            }
            QRadioButton, QCheckBox {
                background: transparent;
                border: none;
            }
        """)

        # Add drop shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(22)
        shadow.setYOffset(3)
        shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(shadow)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(16, 12, 16, 12)
        layout.setSpacing(0)

        # Fixed header with title and zoom controls
        header_widget = QWidget()
        header_widget.setStyleSheet("background: transparent;")
        header_layout = QVBoxLayout(header_widget)
        header_layout.setContentsMargins(0, 0, 0, 6)
        header_layout.setSpacing(4)

        # Title row with zoom controls
        title_row = QHBoxLayout()
        title_row.setContentsMargins(0, 0, 0, 0)
        title_row.setSpacing(4)

        self.title_label = QLabel(title)
        self.title_label.setObjectName("cardTitle")
        self.title_label.setFixedHeight(24)
        title_row.addWidget(self.title_label)
        title_row.addStretch()

        header_layout.addLayout(title_row)

        # Divider
        divider = QFrame()
        divider.setFixedHeight(1)
        divider.setStyleSheet("background: #e5e7eb;")
        header_layout.addWidget(divider)

        layout.addWidget(header_widget)

        # Scrollable content area
        self.content_scroll = QScrollArea()
        self.content_scroll.setWidgetResizable(True)
        self.content_scroll.setFrameShape(QFrame.Shape.NoFrame)
        self.content_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.content_scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.content_scroll.setStyleSheet("QScrollArea { background: transparent; border: none; }")

        self.content_widget = QWidget()
        self.content_widget.setStyleSheet("background: transparent;")
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(0, 6, 0, 0)
        self.content_layout.setSpacing(2)

        self.preview_label = MyPsychAdminRichTextEditor()
        self.preview_label.setPlaceholderText("Click to edit...")
        self.preview_label.setObjectName("cardContent")
        self.preview_label.setStyleSheet("""
            QTextEdit {
                font-size: 20px;
                color: #374151;
                padding: 4px;
                background: transparent;
                border: none;
            }
        """)
        self.preview_label.setFrameShape(QFrame.Shape.NoFrame)
        self.content_layout.addWidget(self.preview_label)
        self.content_layout.addStretch()

        self.content_scroll.setWidget(self.content_widget)
        layout.addWidget(self.content_scroll, 1)

        # Add zoom controls to title row
        zoom_row = create_zoom_row(self.preview_label, base_size=12)
        for i in range(zoom_row.count()):
            item = zoom_row.itemAt(i)
            if item.widget():
                title_row.addWidget(item.widget())

    def set_preview(self, text: str):
        self.preview_label.setPlainText(text)

    def get_content(self) -> str:
        return self.preview_label.toPlainText()

    def mousePressEvent(self, event):
        self.clicked.emit(self.key)
        super().mousePressEvent(event)


# ================================================================
# TOOLBAR
# ================================================================
class A6Toolbar(QWidget):
    """Toolbar for the A6 Form Page - same as LetterToolbar."""

    export_docx = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setFixedHeight(80)
        self.setStyleSheet("""
            A6Toolbar {
                background: rgba(200, 215, 220, 0.95);
                border-bottom: 1px solid rgba(0,0,0,0.12);
            }
            QToolButton {
                background: transparent;
                color: #333333;
                padding: 6px 10px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: 500;
            }
            QToolButton:hover { background: rgba(0,0,0,0.08); }
            QScrollArea { background: transparent; border: none; }
            QScrollBar:horizontal { height: 6px; background: transparent; }
            QScrollBar::handle:horizontal { background: rgba(0,0,0,0.2); border-radius: 3px; min-width: 30px; }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal { width: 0px; }
        """)

        outer_layout = QHBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        outer_layout.setSpacing(0)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setFrameShape(QScrollArea.Shape.NoFrame)
        scroll.setFixedHeight(80)

        container = QWidget()
        container.setStyleSheet("background: transparent;")
        container.setFixedHeight(76)
        layout = QHBoxLayout(container)
        layout.setContentsMargins(8, 2, 24, 2)
        layout.setSpacing(10)

        # Export DOCX button
        export_btn = QToolButton()
        export_btn.setText("Export DOCX")
        export_btn.setFixedSize(140, 38)
        export_btn.setStyleSheet("""
            QToolButton { background: #2563eb; color: white; font-size: 13px; font-weight: 600; border: none; border-radius: 8px; }
            QToolButton:hover { background: #1d4ed8; }
        """)
        export_btn.clicked.connect(self.export_docx.emit)
        layout.addStretch()
        layout.addWidget(export_btn)

        scroll.setWidget(container)
        outer_layout.addWidget(scroll)


# ================================================================
# MAIN A6 FORM PAGE
# ================================================================
class A6FormPage(QWidget):
    """Page for completing MHA Form A6 - Section 3 AMHP Application for Treatment."""

    go_back = Signal()

    def __init__(self, db=None, parent=None):
        super().__init__(parent)
        self.db = db
        self._my_details = self._load_my_details()

        self.cards = {}
        self.card_sections = {}

        self.sections = [
            ("Hospital", "hospital"),
            ("AMHP & Patient", "amhp_patient"),
            ("Local Authority", "authority"),
            ("Nearest Relative", "nr"),
            ("Interview & Signature", "interview"),
        ]

        self._setup_ui()
        self._prefill_amhp_details()

    def _load_my_details(self) -> dict:
        if not self.db:
            return {}
        details = self.db.get_clinician_details()
        if not details:
            return {}
        return {
            "full_name": details[1] or "",
            "email": details[7] or "",
        }

    def _prefill_amhp_details(self):
        if self._my_details.get("full_name"):
            self.amhp_name.setText(self._my_details["full_name"])
        if self._my_details.get("email"):
            self.amhp_email.setText(self._my_details["email"])

    def _setup_ui(self):
        main_layout = QVBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)

        # Header
        header = QFrame()
        header.setFixedHeight(60)
        header.setStyleSheet("background: #2563eb; border-bottom: 1px solid #1d4ed8;")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(24, 0, 24, 0)

        back_btn = QPushButton("< Back")
        back_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        back_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.15);
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-size: 17px;
                font-weight: 600;
            }
            QPushButton:hover { background: rgba(255,255,255,0.25); }
        """)
        back_btn.clicked.connect(self._go_back)
        header_layout.addWidget(back_btn)

        title = QLabel("Form A6 — Section 3 Application by AMHP")
        title.setStyleSheet("font-size: 22px; font-weight: 700; color: white;")
        header_layout.addWidget(title)
        header_layout.addStretch()

        clear_btn = QPushButton("Clear Form")
        clear_btn.setFixedSize(220, 36)
        clear_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: #991b1b;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 6px;
                font-weight: 600;
                font-size: 17px;
            }
            QPushButton:hover { background: #7f1d1d; }
            QPushButton:pressed { background: #450a0a; }
        """)
        clear_btn.clicked.connect(self._clear_form)
        header_layout.addWidget(clear_btn)

        main_layout.addWidget(header)

        # Toolbar
        self.toolbar = MHAFormToolbar()
        self.toolbar.export_docx.connect(self._export_docx)

        self._active_editor = None

        def get_active_editor():
            return self._active_editor

        def safe_call(method_name):
            editor = get_active_editor()
            if editor and hasattr(editor, method_name):
                getattr(editor, method_name)()

        self.toolbar.set_font_family.connect(
            lambda family: get_active_editor().set_font_family(family) if get_active_editor() else None
        )
        self.toolbar.set_font_size.connect(
            lambda size: get_active_editor().set_font_size(size) if get_active_editor() else None
        )
        self.toolbar.toggle_bold.connect(lambda: safe_call("toggle_bold"))
        self.toolbar.toggle_italic.connect(lambda: safe_call("toggle_italic"))
        self.toolbar.toggle_underline.connect(lambda: safe_call("toggle_underline"))
        self.toolbar.set_text_color.connect(
            lambda c: get_active_editor().set_text_color(c) if get_active_editor() else None
        )
        self.toolbar.set_highlight_color.connect(
            lambda c: get_active_editor().set_highlight_color(c) if get_active_editor() else None
        )
        self.toolbar.set_align_left.connect(lambda: safe_call("align_left"))
        self.toolbar.set_align_center.connect(lambda: safe_call("align_center"))
        self.toolbar.set_align_right.connect(lambda: safe_call("align_right"))
        self.toolbar.set_align_justify.connect(lambda: safe_call("align_justify"))
        self.toolbar.bullet_list.connect(lambda: safe_call("bullet_list"))
        self.toolbar.numbered_list.connect(lambda: safe_call("numbered_list"))
        self.toolbar.indent.connect(lambda: safe_call("indent"))
        self.toolbar.outdent.connect(lambda: safe_call("outdent"))
        self.toolbar.undo.connect(lambda: safe_call("undo"))
        self.toolbar.redo.connect(lambda: safe_call("redo"))
        self.toolbar.insert_date.connect(lambda: safe_call("insert_date"))
        self.toolbar.insert_section_break.connect(lambda: safe_call("insert_section_break"))

        def check_spelling():
            editor = get_active_editor()
            if editor and hasattr(editor, 'jump_to_next_error'):
                if not editor.jump_to_next_error():
                    QMessageBox.information(self, "Spell Check", "No spelling errors found.")

        self.toolbar.check_spelling.connect(check_spelling)

        main_layout.addWidget(self.toolbar)

        # Splitter: Cards | Popup
        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setHandleWidth(6)
        splitter.setStyleSheet("QSplitter::handle { background: #d1d5db; } QSplitter::handle:hover { background: #6BAF8D; }")

        # Left: Cards
        cards_scroll = QScrollArea()
        cards_scroll.setWidgetResizable(True)
        cards_scroll.setFrameShape(QFrame.Shape.NoFrame)
        cards_scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        cards_scroll.setStyleSheet("background: #f9fafb;")

        cards_container = QWidget()
        cards_container.setStyleSheet("background: #f9fafb;")
        cards_layout = QVBoxLayout(cards_container)
        cards_layout.setContentsMargins(16, 16, 16, 16)
        cards_layout.setSpacing(8)

        for title, key in self.sections:
            section = ResizableSection()
            section.set_content_height(180)
            section._min_height = 120
            section._max_height = 350

            card = A6CardWidget(title, key)
            card.clicked.connect(self._on_card_clicked)
            self._hook_editor_focus(card.preview_label)
            self.cards[key] = card

            section.set_content(card)
            self.card_sections[key] = section
            cards_layout.addWidget(section)

        cards_layout.addStretch()
        cards_scroll.setWidget(cards_container)
        splitter.addWidget(cards_scroll)

        # Right: Popup panel
        self.popup_panel = QFrame()
        self.popup_panel.setMinimumWidth(400)
        self.popup_panel.setMaximumWidth(650)
        self.popup_panel.setStyleSheet("background: white; border-left: 1px solid #e5e7eb;")

        popup_layout = QVBoxLayout(self.popup_panel)
        popup_layout.setContentsMargins(0, 0, 0, 0)
        popup_layout.setSpacing(0)

        self.popup_stack = QStackedWidget()
        self._build_all_popups()
        popup_layout.addWidget(self.popup_stack)

        # Initialize cards with default date values
        self._update_interview_card()

        splitter.addWidget(self.popup_panel)
        splitter.setSizes([400, 520])

        main_layout.addWidget(splitter, 1)

        # Show first popup by default
        self._on_card_clicked("hospital")

    def _on_card_clicked(self, key: str):
        index_map = {k: i for i, (_, k) in enumerate(self.sections)}
        if key in index_map:
            self.popup_stack.setCurrentIndex(index_map[key])

    def _build_all_popups(self):
        self._build_hospital_popup()
        self._build_amhp_patient_popup()
        self._build_authority_popup()
        self._build_nr_popup()
        self._build_interview_popup()

    def _create_popup_scroll(self) -> tuple:
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setFrameShape(QFrame.Shape.NoFrame)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setStyleSheet("QScrollArea { background: white; border: none; }")

        container = QWidget()
        container.setMaximumWidth(500)
        container.setStyleSheet("background: white;")
        layout = QVBoxLayout(container)
        layout.setContentsMargins(20, 16, 20, 16)
        layout.setSpacing(12)

        return scroll, container, layout

    def _create_line_edit(self, placeholder: str = "") -> QLineEdit:
        edit = QLineEdit()
        edit.setPlaceholderText(placeholder)
        edit.setStyleSheet("""
            QLineEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 20px;
            }
            QLineEdit:focus { border-color: #2563eb; }
        """)
        return edit

    def _create_date_edit(self) -> QDateEdit:
        date_edit = QDateEdit()
        date_edit.setCalendarPopup(True)
        date_edit.setDate(QDate.currentDate())
        date_edit.setStyleSheet("""
            QDateEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 8px 10px;
                font-size: 20px;
            }
            QDateEdit::drop-down { border: none; width: 20px; }
        """)
        return date_edit

    # ----------------------------------------------------------------
    # HOSPITAL POPUP
    # ----------------------------------------------------------------
    def _build_hospital_popup(self):
        scroll, container, layout = self._create_popup_scroll()

        header = QLabel("Hospital Details")
        header.setStyleSheet("font-size: 19px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Hospital frame
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 6px; }")
        frame_layout = QVBoxLayout(frame)
        frame_layout.setContentsMargins(12, 10, 12, 10)
        frame_layout.setSpacing(8)

        name_lbl = QLabel("Hospital Name:")
        name_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        frame_layout.addWidget(name_lbl)
        self.hospital_name = self._create_line_edit("Hospital name")
        frame_layout.addWidget(self.hospital_name)

        addr_lbl = QLabel("Address:")
        addr_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        frame_layout.addWidget(addr_lbl)
        self.hospital_address = self._create_line_edit("Hospital address")
        frame_layout.addWidget(self.hospital_address)

        layout.addWidget(frame)

        # Auto-sync to card
        self.hospital_name.textChanged.connect(self._update_hospital_card)
        self.hospital_address.textChanged.connect(self._update_hospital_card)

        layout.addStretch()

        scroll.setWidget(container)
        self.popup_stack.addWidget(scroll)

    def _update_hospital_card(self):
        parts = []
        if self.hospital_name.text():
            parts.append(self.hospital_name.text())
        if self.hospital_address.text():
            parts.append(self.hospital_address.text()[:40])
        self.cards["hospital"].set_preview("\n".join(parts) if parts else "No data entered")

    # ----------------------------------------------------------------
    # AMHP & PATIENT POPUP
    # ----------------------------------------------------------------
    def _build_amhp_patient_popup(self):
        scroll, container, layout = self._create_popup_scroll()

        header = QLabel("AMHP & Patient Details")
        header.setStyleSheet("font-size: 19px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # AMHP frame
        amhp_frame = QFrame()
        amhp_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 6px; }")
        amhp_layout = QVBoxLayout(amhp_frame)
        amhp_layout.setContentsMargins(12, 10, 12, 10)
        amhp_layout.setSpacing(6)

        amhp_header = QLabel("AMHP Details")
        amhp_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #166534;")
        amhp_layout.addWidget(amhp_header)

        self.amhp_name = self._create_line_edit("Full name")
        amhp_layout.addWidget(self.amhp_name)

        self.amhp_address = self._create_line_edit("Address")
        amhp_layout.addWidget(self.amhp_address)

        self.amhp_email = self._create_line_edit("Email")
        amhp_layout.addWidget(self.amhp_email)

        layout.addWidget(amhp_frame)

        # Patient frame
        patient_frame = QFrame()
        patient_frame.setStyleSheet("QFrame { background: #faf5ff; border: none; border-radius: 6px; }")
        patient_layout = QVBoxLayout(patient_frame)
        patient_layout.setContentsMargins(12, 10, 12, 10)
        patient_layout.setSpacing(6)

        patient_header = QLabel("Patient Details")
        patient_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #7c3aed;")
        patient_layout.addWidget(patient_header)

        self.patient_name = self._create_line_edit("Patient full name")
        patient_layout.addWidget(self.patient_name)

        self.patient_address = self._create_line_edit("Patient address")
        patient_layout.addWidget(self.patient_address)

        layout.addWidget(patient_frame)

        # Auto-sync to card
        self.amhp_name.textChanged.connect(self._update_amhp_patient_card)
        self.amhp_address.textChanged.connect(self._update_amhp_patient_card)
        self.amhp_email.textChanged.connect(self._update_amhp_patient_card)
        self.patient_name.textChanged.connect(self._update_amhp_patient_card)
        self.patient_address.textChanged.connect(self._update_amhp_patient_card)

        layout.addStretch()

        scroll.setWidget(container)
        self.popup_stack.addWidget(scroll)

    def _update_amhp_patient_card(self):
        lines = []
        if self.amhp_name.text():
            lines.append(self.amhp_name.text())
        if self.amhp_address.text():
            lines.append(self.amhp_address.text())
        if self.patient_name.text():
            if lines:
                lines.append("")  # Blank line separator
            lines.append(self.patient_name.text())
        if self.patient_address.text():
            lines.append(self.patient_address.text())
        self.cards["amhp_patient"].set_preview("\n".join(lines) if lines else "No data entered")

    # ----------------------------------------------------------------
    # LOCAL AUTHORITY POPUP
    # ----------------------------------------------------------------
    def _build_authority_popup(self):
        scroll, container, layout = self._create_popup_scroll()

        header = QLabel("Local Authority")
        header.setStyleSheet("font-size: 19px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Authority frame
        frame = QFrame()
        frame.setStyleSheet("QFrame { background: #fefce8; border: none; border-radius: 6px; }")
        frame_layout = QVBoxLayout(frame)
        frame_layout.setContentsMargins(12, 10, 12, 10)
        frame_layout.setSpacing(8)

        acting_lbl = QLabel("Acting on behalf of:")
        acting_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        frame_layout.addWidget(acting_lbl)
        self.local_authority = self._create_line_edit("Local social services authority")
        frame_layout.addWidget(self.local_authority)

        # Approved by section
        approved_lbl = QLabel("Approved by:")
        approved_lbl.setStyleSheet("font-size: 20px; font-weight: 600; color: #374151; margin-top: 8px;")
        frame_layout.addWidget(approved_lbl)

        radio_row = QHBoxLayout()
        self.approved_by_same = QRadioButton("Same authority")
        self.approved_by_different = QRadioButton("Different authority")
        self.approved_by_same.setChecked(True)
        self.approved_btn_group = QButtonGroup(self)
        self.approved_btn_group.addButton(self.approved_by_same)
        self.approved_btn_group.addButton(self.approved_by_different)
        radio_row.addWidget(self.approved_by_same)
        radio_row.addWidget(self.approved_by_different)
        radio_row.addStretch()
        frame_layout.addLayout(radio_row)

        self.approved_by_authority = self._create_line_edit("Approving authority if different")
        self.approved_by_authority.setEnabled(False)
        self.approved_by_different.toggled.connect(self.approved_by_authority.setEnabled)
        frame_layout.addWidget(self.approved_by_authority)

        layout.addWidget(frame)

        # Auto-sync to card
        self.local_authority.textChanged.connect(self._update_authority_card)
        self.approved_by_same.toggled.connect(self._update_authority_card)
        self.approved_by_authority.textChanged.connect(self._update_authority_card)

        layout.addStretch()

        scroll.setWidget(container)
        self.popup_stack.addWidget(scroll)

    def _update_authority_card(self):
        parts = []
        if self.local_authority.text():
            parts.append(self.local_authority.text())
        if self.approved_by_different.isChecked() and self.approved_by_authority.text():
            parts.append(self.approved_by_authority.text())
        self.cards["authority"].set_preview("\n".join(parts) if parts else "No data entered")

    # ----------------------------------------------------------------
    # NEAREST RELATIVE POPUP
    # ----------------------------------------------------------------
    def _build_nr_popup(self):
        scroll, container, layout = self._create_popup_scroll()

        header = QLabel("Nearest Relative Consultation")
        header.setStyleSheet("font-size: 19px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Main choice
        main_row = QHBoxLayout()
        self.nr_consulted = QRadioButton("Consulted")
        self.nr_not_consulted = QRadioButton("NOT Consulted")
        self.nr_consulted.setChecked(True)
        self.main_nr_group = QButtonGroup(self)
        self.main_nr_group.addButton(self.nr_consulted)
        self.main_nr_group.addButton(self.nr_not_consulted)
        self.nr_consulted.setStyleSheet("font-size: 20px; font-weight: 600;")
        self.nr_not_consulted.setStyleSheet("font-size: 20px; font-weight: 600;")
        main_row.addWidget(self.nr_consulted)
        main_row.addWidget(self.nr_not_consulted)
        main_row.addStretch()
        layout.addLayout(main_row)

        # === CONSULTED SECTION ===
        self.consulted_widget = QFrame()
        self.consulted_widget.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 6px; }")
        consulted_layout = QVBoxLayout(self.consulted_widget)
        consulted_layout.setContentsMargins(12, 10, 12, 10)
        consulted_layout.setSpacing(8)

        opt_row = QHBoxLayout()
        self.nr_option_a = QRadioButton("(a) Is NR")
        self.nr_option_b = QRadioButton("(b) Authorised by court/NR")
        self.nr_option_a.setChecked(True)
        self.option_btn_group = QButtonGroup(self)
        self.option_btn_group.addButton(self.nr_option_a)
        self.option_btn_group.addButton(self.nr_option_b)
        opt_row.addWidget(self.nr_option_a)
        opt_row.addWidget(self.nr_option_b)
        opt_row.addStretch()
        consulted_layout.addLayout(opt_row)

        self.nr_name = self._create_line_edit("NR full name")
        consulted_layout.addWidget(self.nr_name)

        self.nr_address = self._create_line_edit("NR address")
        consulted_layout.addWidget(self.nr_address)

        layout.addWidget(self.consulted_widget)

        # === NOT CONSULTED SECTION ===
        self.not_consulted_widget = QFrame()
        self.not_consulted_widget.setStyleSheet("QFrame { background: #fef2f2; border: none; border-radius: 6px; }")
        not_consulted_layout = QVBoxLayout(self.not_consulted_widget)
        not_consulted_layout.setContentsMargins(12, 10, 12, 10)
        not_consulted_layout.setSpacing(8)

        nc_opt_row = QHBoxLayout()
        self.nc_option_a = QRadioButton("(a) Unable to ascertain")
        self.nc_option_b = QRadioButton("(b) No NR")
        self.nc_option_c = QRadioButton("(c) Not practicable")
        self.nc_option_a.setChecked(True)
        self.nc_option_group = QButtonGroup(self)
        self.nc_option_group.addButton(self.nc_option_a)
        self.nc_option_group.addButton(self.nc_option_b)
        self.nc_option_group.addButton(self.nc_option_c)
        nc_opt_row.addWidget(self.nc_option_a)
        nc_opt_row.addWidget(self.nc_option_b)
        nc_opt_row.addWidget(self.nc_option_c)
        not_consulted_layout.addLayout(nc_opt_row)

        # Option C details
        self.nc_option_c_widget = QWidget()
        nc_c_layout = QVBoxLayout(self.nc_option_c_widget)
        nc_c_layout.setContentsMargins(0, 8, 0, 0)
        nc_c_layout.setSpacing(6)

        self.nc_nr_name = self._create_line_edit("NR full name")
        nc_c_layout.addWidget(self.nc_nr_name)

        self.nc_nr_address = self._create_line_edit("NR address")
        nc_c_layout.addWidget(self.nc_nr_address)

        nc_sub_row = QHBoxLayout()
        self.nc_c_is_nr = QRadioButton("Is NR")
        self.nc_c_is_authorised = QRadioButton("Authorised")
        self.nc_c_is_nr.setChecked(True)
        self.nc_c_sub_group = QButtonGroup(self)
        self.nc_c_sub_group.addButton(self.nc_c_is_nr)
        self.nc_c_sub_group.addButton(self.nc_c_is_authorised)
        nc_sub_row.addWidget(self.nc_c_is_nr)
        nc_sub_row.addWidget(self.nc_c_is_authorised)
        nc_sub_row.addStretch()
        nc_c_layout.addLayout(nc_sub_row)

        reason_lbl = QLabel("Reason:")
        reason_lbl.setStyleSheet("font-size: 19px; font-weight: 500; color: #374151;")
        nc_c_layout.addWidget(reason_lbl)
        self.nc_reason = QTextEdit()
        self.nc_reason.setPlaceholderText("Reason why not practicable / unreasonable delay")
        self.nc_reason.setMaximumHeight(60)
        self.nc_reason.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 6px;
                font-size: 19px;
            }
        """)
        nc_c_layout.addWidget(self.nc_reason)

        not_consulted_layout.addWidget(self.nc_option_c_widget)
        self.nc_option_c_widget.hide()

        layout.addWidget(self.not_consulted_widget)
        self.not_consulted_widget.hide()

        # Connect visibility toggles
        self.nr_consulted.toggled.connect(self._toggle_nr_sections)
        self.nc_option_c.toggled.connect(self.nc_option_c_widget.setVisible)

        # Auto-sync to card
        self.nr_consulted.toggled.connect(self._update_nr_card)
        self.nr_name.textChanged.connect(self._update_nr_card)
        self.nr_address.textChanged.connect(self._update_nr_card)
        self.nr_option_a.toggled.connect(self._update_nr_card)
        self.nc_option_a.toggled.connect(self._update_nr_card)
        self.nc_option_b.toggled.connect(self._update_nr_card)
        self.nc_option_c.toggled.connect(self._update_nr_card)
        self.nc_nr_name.textChanged.connect(self._update_nr_card)
        self.nc_nr_address.textChanged.connect(self._update_nr_card)

        layout.addStretch()

        scroll.setWidget(container)
        self.popup_stack.addWidget(scroll)

    def _toggle_nr_sections(self, consulted: bool):
        self.consulted_widget.setVisible(consulted)
        self.not_consulted_widget.setVisible(not consulted)

    def _update_nr_card(self):
        parts = []
        if self.nr_consulted.isChecked():
            if self.nr_name.text():
                parts.append(self.nr_name.text())
            if self.nr_address.text():
                parts.append(self.nr_address.text())
        else:
            # Not consulted - show NR details if option C (not practicable)
            if self.nc_option_c.isChecked():
                if self.nc_nr_name.text():
                    parts.append(self.nc_nr_name.text())
                if self.nc_nr_address.text():
                    parts.append(self.nc_nr_address.text())
        self.cards["nr"].set_preview("\n".join(parts) if parts else "No data entered")

    # ----------------------------------------------------------------
    # INTERVIEW & SIGNATURE POPUP
    # ----------------------------------------------------------------
    def _build_interview_popup(self):
        scroll, container, layout = self._create_popup_scroll()

        header = QLabel("Interview & Signature")
        header.setStyleSheet("font-size: 19px; font-weight: 700; color: #1f2937;")
        layout.addWidget(header)

        # Interview frame
        interview_frame = QFrame()
        interview_frame.setStyleSheet("QFrame { background: #eff6ff; border: none; border-radius: 6px; }")
        interview_layout = QVBoxLayout(interview_frame)
        interview_layout.setContentsMargins(12, 10, 12, 10)
        interview_layout.setSpacing(8)

        interview_header = QLabel("Patient Interview")
        interview_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #1e40af;")
        interview_layout.addWidget(interview_header)

        date_row = QHBoxLayout()
        date_lbl = QLabel("Date seen:")
        date_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        date_row.addWidget(date_lbl)
        self.last_seen_date = self._create_date_edit()
        self.last_seen_date.setFixedWidth(140)
        date_row.addWidget(self.last_seen_date)
        info = QLabel("(within 14 days)")
        info.setStyleSheet("font-size: 14px; color: #6b7280;")
        date_row.addWidget(info)
        date_row.addStretch()
        interview_layout.addLayout(date_row)

        layout.addWidget(interview_frame)

        # Medical recs frame
        med_frame = QFrame()
        med_frame.setStyleSheet("QFrame { background: #fefce8; border: none; border-radius: 6px; }")
        med_layout = QVBoxLayout(med_frame)
        med_layout.setContentsMargins(12, 10, 12, 10)
        med_layout.setSpacing(6)

        med_header = QLabel("Medical Recommendations")
        med_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #854d0e;")
        med_layout.addWidget(med_header)

        acq_lbl = QLabel("If neither practitioner had previous acquaintance:")
        acq_lbl.setStyleSheet("font-size: 19px; color: #374151;")
        med_layout.addWidget(acq_lbl)

        self.no_acquaintance_reason = QTextEdit()
        self.no_acquaintance_reason.setPlaceholderText("Explain why (leave blank if N/A)")
        self.no_acquaintance_reason.setMaximumHeight(60)
        self.no_acquaintance_reason.setStyleSheet("""
            QTextEdit {
                background: white;
                border: 1px solid #d1d5db;
                border-radius: 6px;
                padding: 6px;
                font-size: 19px;
            }
        """)
        med_layout.addWidget(self.no_acquaintance_reason)

        layout.addWidget(med_frame)

        # Signature frame
        sig_frame = QFrame()
        sig_frame.setStyleSheet("QFrame { background: #f0fdf4; border: none; border-radius: 6px; }")
        sig_layout = QVBoxLayout(sig_frame)
        sig_layout.setContentsMargins(12, 10, 12, 10)
        sig_layout.setSpacing(8)

        sig_header = QLabel("Signature")
        sig_header.setStyleSheet("font-size: 20px; font-weight: 700; color: #166534;")
        sig_layout.addWidget(sig_header)

        sig_row = QHBoxLayout()
        sig_lbl = QLabel("Date:")
        sig_lbl.setStyleSheet("font-size: 20px; font-weight: 500; color: #374151;")
        sig_row.addWidget(sig_lbl)
        self.signature_date = self._create_date_edit()
        self.signature_date.setFixedWidth(140)
        sig_row.addWidget(self.signature_date)
        sig_row.addStretch()
        sig_layout.addLayout(sig_row)

        layout.addWidget(sig_frame)

        # Auto-sync to card
        self.last_seen_date.dateChanged.connect(self._update_interview_card)
        self.no_acquaintance_reason.textChanged.connect(self._update_interview_card)
        self.signature_date.dateChanged.connect(self._update_interview_card)

        layout.addStretch()

        scroll.setWidget(container)
        self.popup_stack.addWidget(scroll)

    def _update_interview_card(self):
        parts = []
        parts.append(self.last_seen_date.date().toString('dd MMM yyyy'))
        parts.append(self.signature_date.date().toString('dd MMM yyyy'))
        self.cards["interview"].set_preview("\n".join(parts))

    def _register_active_editor(self, editor):
        """Register an editor as the active editor for toolbar actions."""
        self._active_editor = editor

    def _hook_editor_focus(self, editor):
        """Hook focus events on an editor to register it as active when clicked."""
        original_focus_in = editor.focusInEvent
        page = self

        def focus_handler(event):
            page._register_active_editor(editor)
            original_focus_in(event)

        editor.focusInEvent = focus_handler

    # ----------------------------------------------------------------
    # Actions
    # ----------------------------------------------------------------
    def _go_back(self):
        self.go_back.emit()

    def _clear_form(self):
        reply = QMessageBox.question(
            self, "Clear Form", "Are you sure you want to clear all form data?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            self.hospital_name.clear()
            self.hospital_address.clear()
            self.amhp_name.clear()
            self.amhp_address.clear()
            self.amhp_email.clear()
            self.patient_name.clear()
            self.patient_address.clear()
            self.local_authority.clear()
            self.approved_by_same.setChecked(True)
            self.approved_by_authority.clear()
            self.nr_consulted.setChecked(True)
            self.nr_option_a.setChecked(True)
            self.nr_name.clear()
            self.nr_address.clear()
            self.nc_option_a.setChecked(True)
            self.nc_nr_name.clear()
            self.nc_nr_address.clear()
            self.nc_c_is_nr.setChecked(True)
            self.nc_reason.clear()
            self.last_seen_date.setDate(QDate.currentDate())
            self.no_acquaintance_reason.clear()
            self.signature_date.setDate(QDate.currentDate())
            # Update cards
            self._update_hospital_card()
            self._update_amhp_patient_card()
            self._update_authority_card()
            self._update_nr_card()
            self._update_interview_card()
            # Restore my details fields
            self._prefill_amhp_details()

    def _export_docx(self):
        """Export the form to DOCX format using the official template."""
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Form A6",
            f"Form_A6_{datetime.now().strftime('%Y%m%d')}.docx",
            "Word Documents (*.docx)"
        )
        if not file_path:
            return

        try:
            import os
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            template_path = resource_path('templates', 'Form_A6_template.docx')

            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Template Missing", "Form A6 template not found.")
                return

            doc = Document(template_path)
            # Remove document protection so exported file is editable
            protection = doc.settings.element.find(qn('w:documentProtection'))
            if protection is not None:
                doc.settings.element.remove(protection)

            # Clean ALL paragraphs - remove permission markers and convert grey to cream
            for para in doc.paragraphs:
                para_xml = para._element
                # Remove permission markers (they cause grey appearance)
                for perm_start in para_xml.findall('.//' + qn('w:permStart')):
                    perm_start.getparent().remove(perm_start)
                for perm_end in para_xml.findall('.//' + qn('w:permEnd')):
                    perm_end.getparent().remove(perm_end)
                # Convert paragraph-level grey shading to cream (keep the shading, just change color)
                pPr = para_xml.find(qn('w:pPr'))
                if pPr is not None:
                    for shd in pPr.findall(qn('w:shd')):
                        shd.set(qn('w:fill'), 'FFFED5')
                # Convert run-level grey shading to cream
                for run in para.runs:
                    rPr = run._element.find(qn('w:rPr'))
                    if rPr is not None:
                        for shd in rPr.findall(qn('w:shd')):
                            shd.set(qn('w:fill'), 'FFFED5')

            def set_para_text(para, new_text):
                """Set paragraph text - if empty, use spaces for visible blank placeholder"""
                if not new_text.strip():
                    new_text = '                                                                   '
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                else:
                    run = para.add_run(new_text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)

            # Gold bracket color - brighter gold #918C0D
            BRACKET_COLOR = RGBColor(0x91, 0x8C, 0x0D)

            def set_entry_box(para, content=""):
                """Set entry box with bold gold brackets [content] - brackets are gold, content is cream"""
                if not content.strip():
                    content = '                                                                   '
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)
                # Add opening bracket with gold color and bold
                if para.runs:
                    para.runs[0].text = '['
                    bracket_open = para.runs[0]
                else:
                    bracket_open = para.add_run('[')
                bracket_open.font.name = 'Arial'
                bracket_open.font.size = Pt(12)
                bracket_open.font.bold = True
                bracket_open.font.color.rgb = BRACKET_COLOR
                rPr = bracket_open._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'FFFED5')
                rPr.append(shd)
                # Add content with cream highlighting
                content_run = para.add_run(content)
                content_run.font.name = 'Arial'
                content_run.font.size = Pt(12)
                rPr2 = content_run._element.get_or_add_rPr()
                shd2 = OxmlElement('w:shd')
                shd2.set(qn('w:val'), 'clear')
                shd2.set(qn('w:color'), 'auto')
                shd2.set(qn('w:fill'), 'FFFED5')
                rPr2.append(shd2)
                # Add closing bracket with gold color and bold
                bracket_close = para.add_run(']')
                bracket_close.font.name = 'Arial'
                bracket_close.font.size = Pt(12)
                bracket_close.font.bold = True
                bracket_close.font.color.rgb = BRACKET_COLOR
                rPr3 = bracket_close._element.get_or_add_rPr()
                shd3 = OxmlElement('w:shd')
                shd3.set(qn('w:val'), 'clear')
                shd3.set(qn('w:color'), 'auto')
                shd3.set(qn('w:fill'), 'FFFED5')
                rPr3.append(shd3)

            def strikethrough_para(para):
                for run in para.runs:
                    run.font.strike = True

            def highlight_yellow(para):
                """Rebuild paragraph with cream highlight - more robust approach"""
                # Get all text from the paragraph
                text = ''.join(run.text for run in para.runs)
                if not text:
                    text = para.text
                if not text.strip():
                    return  # Nothing to highlight

                # Clear existing runs
                for run in para.runs:
                    run.text = ""
                while len(para.runs) > 1:
                    para._element.remove(para.runs[-1]._element)

                # Add text with cream highlight
                if para.runs:
                    para.runs[0].text = text
                    para.runs[0].font.name = 'Arial'
                    para.runs[0].font.size = Pt(12)
                    rPr = para.runs[0]._element.get_or_add_rPr()
                    # Remove any existing shading first
                    for old_shd in rPr.findall(qn('w:shd')):
                        rPr.remove(old_shd)
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)
                else:
                    run = para.add_run(text)
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    rPr = run._element.get_or_add_rPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'FFFED5')
                    rPr.append(shd)

            paragraphs = doc.paragraphs

            # Hospital (use card content) - gold bracketed entry box
            hospital_text = self.cards["hospital"].get_content().replace("\n", ", ")
            if hospital_text.strip() and hospital_text != "No data entered":
                set_entry_box(paragraphs[3], hospital_text)
            else:
                set_entry_box(paragraphs[3], "")

            # AMHP (use field values with email) - gold bracketed entry box
            amhp_parts = []
            if self.amhp_name.text().strip():
                amhp_parts.append(self.amhp_name.text().strip())
            if self.amhp_address.text().strip():
                amhp_parts.append(self.amhp_address.text().strip())
            if self.amhp_email.text().strip():
                amhp_parts.append(self.amhp_email.text().strip())
            if amhp_parts:
                set_entry_box(paragraphs[5], ", ".join(amhp_parts))
            else:
                set_entry_box(paragraphs[5], "")

            # Patient - gold bracketed entry box
            patient_parts = []
            if self.patient_name.text().strip():
                patient_parts.append(self.patient_name.text().strip())
            if self.patient_address.text().strip():
                patient_parts.append(self.patient_address.text().strip())
            if patient_parts:
                set_entry_box(paragraphs[7], ", ".join(patient_parts))
            else:
                set_entry_box(paragraphs[7], "")

            # Local authority - gold bracketed entry box
            authority_text = self.cards["authority"].get_content().replace("\n", ", ")
            if authority_text.strip() and authority_text != "No data entered":
                set_entry_box(paragraphs[10], authority_text)
            else:
                set_entry_box(paragraphs[10], "")

            # Approved by - "that authority" with gold brackets
            set_entry_box(paragraphs[12], "that authority")
            # Different authority entry box
            if self.approved_by_same.isChecked():
                strikethrough_para(paragraphs[13])
                set_entry_box(paragraphs[14], "")
                strikethrough_para(paragraphs[14])
            else:
                strikethrough_para(paragraphs[12])
                if self.approved_by_authority.text():
                    set_entry_box(paragraphs[14], self.approved_by_authority.text())
                else:
                    set_entry_box(paragraphs[14], "")

            # NR consultation section - wrap in gold brackets with highlighting
            # Para 17: (a) I have consulted... - gold bracket at start
            # Para 18: NR entry box
            # Para 19: who to the best of my knowledge...
            # Para 20: (b) I have consulted...
            # Para 21: NR entry box
            # Para 22: who I understand has been authorised...
            # Para 23: That person has not notified me... - ends with gold bracket

            # Debug: Always show which NR option is selected
            print(f"DEBUG: NR CONSULTED = {self.nr_consulted.isChecked()}, NR NOT CONSULTED = {self.nr_not_consulted.isChecked()}")

            if self.nr_consulted.isChecked():
                nr_text = self.cards["nr"].get_content().replace("\n", ", ")

                # Para 17: (a) I have consulted... - rebuild with gold bracket at start and cream highlight
                p17 = paragraphs[17]
                p17_text = ''.join(run.text for run in p17.runs)
                for run in p17.runs:
                    run.text = ""
                while len(p17.runs) > 1:
                    p17._element.remove(p17.runs[-1]._element)
                # Add gold opening bracket
                if p17.runs:
                    p17.runs[0].text = '['
                    p17.runs[0].font.bold = True
                    p17.runs[0].font.color.rgb = BRACKET_COLOR
                    rPr17o = p17.runs[0]._element.get_or_add_rPr()
                    shd17o = OxmlElement('w:shd')
                    shd17o.set(qn('w:val'), 'clear')
                    shd17o.set(qn('w:color'), 'auto')
                    shd17o.set(qn('w:fill'), 'FFFED5')
                    rPr17o.append(shd17o)
                # Add text with cream highlight
                p17_content = p17.add_run(p17_text)
                p17_content.font.name = 'Arial'
                p17_content.font.size = Pt(12)
                rPr17c = p17_content._element.get_or_add_rPr()
                shd17c = OxmlElement('w:shd')
                shd17c.set(qn('w:val'), 'clear')
                shd17c.set(qn('w:color'), 'auto')
                shd17c.set(qn('w:fill'), 'FFFED5')
                rPr17c.append(shd17c)

                # Para 19: who to the best of my knowledge... - cream highlight
                highlight_yellow(paragraphs[19])

                # Para 20: (b) I have consulted... - cream highlight
                highlight_yellow(paragraphs[20])

                # Para 22: who I understand has been authorised... - cream highlight
                highlight_yellow(paragraphs[22])

                # Para 23: That person has not notified me... - cream highlight with gold bracket at end
                p23 = paragraphs[23]
                p23_text = ''.join(run.text for run in p23.runs)
                for run in p23.runs:
                    run.text = ""
                while len(p23.runs) > 1:
                    p23._element.remove(p23.runs[-1]._element)
                if p23.runs:
                    p23.runs[0].text = p23_text
                    p23.runs[0].font.name = 'Arial'
                    p23.runs[0].font.size = Pt(12)
                    rPr23 = p23.runs[0]._element.get_or_add_rPr()
                    shd23 = OxmlElement('w:shd')
                    shd23.set(qn('w:val'), 'clear')
                    shd23.set(qn('w:color'), 'auto')
                    shd23.set(qn('w:fill'), 'FFFED5')
                    rPr23.append(shd23)
                # Add gold closing bracket
                p23_close = p23.add_run(']')
                p23_close.font.bold = True
                p23_close.font.color.rgb = BRACKET_COLOR
                rPr23c = p23_close._element.get_or_add_rPr()
                shd23c = OxmlElement('w:shd')
                shd23c.set(qn('w:val'), 'clear')
                shd23c.set(qn('w:color'), 'auto')
                shd23c.set(qn('w:fill'), 'FFFED5')
                rPr23c.append(shd23c)

                if self.nr_option_a.isChecked():
                    if nr_text.strip():
                        set_entry_box(paragraphs[18], nr_text)
                    else:
                        set_entry_box(paragraphs[18], "")
                    set_entry_box(paragraphs[21], "")
                    strikethrough_para(paragraphs[20])
                    strikethrough_para(paragraphs[21])
                    strikethrough_para(paragraphs[22])
                else:
                    if nr_text.strip():
                        set_entry_box(paragraphs[21], nr_text)
                    else:
                        set_entry_box(paragraphs[21], "")
                    set_entry_box(paragraphs[18], "")
                    strikethrough_para(paragraphs[17])
                    strikethrough_para(paragraphs[18])
                    strikethrough_para(paragraphs[19])

                # Apply cream highlighting to NR NOT CONSULTED section FIRST, then strikethrough
                # Para 25: (a) I have been unable to ascertain...
                # Para 26: (b) To the best of my knowledge and belief...
                # Para 27: (c) I understand that...
                # Para 28: entry box
                # Para 29: is
                # Para 30: (i) this patient's nearest relative...
                # Para 31: (ii) authorised to exercise...
                # Para 33: but in my opinion it is not reasonably practicable...
                # Para 34: entry box
                highlight_yellow(paragraphs[25])
                highlight_yellow(paragraphs[26])
                highlight_yellow(paragraphs[27])
                set_entry_box(paragraphs[28], "")  # blank entry box
                highlight_yellow(paragraphs[29])
                highlight_yellow(paragraphs[30])
                highlight_yellow(paragraphs[31])

                # Para 33: Format with gold brackets around "not reasonably practicable/would involve unreasonable delay"
                p33 = paragraphs[33]
                p33_text = ''.join(run.text for run in p33.runs)
                for run in p33.runs:
                    run.text = ""
                while len(p33.runs) > 1:
                    p33._element.remove(p33.runs[-1]._element)

                # Find the phrase to bracket
                if "not reasonably practicable" in p33_text:
                    idx = p33_text.index("not reasonably practicable")
                    before = p33_text[:idx]
                    end_phrase = "unreasonable delay"
                    if end_phrase in p33_text:
                        end_idx = p33_text.index(end_phrase) + len(end_phrase)
                        bracketed = p33_text[idx:end_idx]
                        after = p33_text[end_idx:]
                    else:
                        bracketed = p33_text[idx:]
                        after = ""
                else:
                    before = p33_text
                    bracketed = ""
                    after = ""

                # Add "before" text with cream highlight
                if p33.runs:
                    p33.runs[0].text = before
                    p33.runs[0].font.name = 'Arial'
                    p33.runs[0].font.size = Pt(12)
                    rPr33b = p33.runs[0]._element.get_or_add_rPr()
                    shd33b = OxmlElement('w:shd')
                    shd33b.set(qn('w:val'), 'clear')
                    shd33b.set(qn('w:color'), 'auto')
                    shd33b.set(qn('w:fill'), 'FFFED5')
                    rPr33b.append(shd33b)

                # Add gold bracketed phrase
                if bracketed:
                    p33_bo = p33.add_run('[')
                    p33_bo.font.bold = True
                    p33_bo.font.color.rgb = BRACKET_COLOR
                    rPr33bo = p33_bo._element.get_or_add_rPr()
                    shd33bo = OxmlElement('w:shd')
                    shd33bo.set(qn('w:val'), 'clear')
                    shd33bo.set(qn('w:color'), 'auto')
                    shd33bo.set(qn('w:fill'), 'FFFED5')
                    rPr33bo.append(shd33bo)

                    p33_bc = p33.add_run(bracketed)
                    p33_bc.font.name = 'Arial'
                    p33_bc.font.size = Pt(12)
                    rPr33bc = p33_bc._element.get_or_add_rPr()
                    shd33bc = OxmlElement('w:shd')
                    shd33bc.set(qn('w:val'), 'clear')
                    shd33bc.set(qn('w:color'), 'auto')
                    shd33bc.set(qn('w:fill'), 'FFFED5')
                    rPr33bc.append(shd33bc)

                    p33_bcl = p33.add_run(']')
                    p33_bcl.font.bold = True
                    p33_bcl.font.color.rgb = BRACKET_COLOR
                    rPr33bcl = p33_bcl._element.get_or_add_rPr()
                    shd33bcl = OxmlElement('w:shd')
                    shd33bcl.set(qn('w:val'), 'clear')
                    shd33bcl.set(qn('w:color'), 'auto')
                    shd33bcl.set(qn('w:fill'), 'FFFED5')
                    rPr33bcl.append(shd33bcl)

                # Add "after" text with cream highlight
                if after:
                    p33_after = p33.add_run(after)
                    p33_after.font.name = 'Arial'
                    p33_after.font.size = Pt(12)
                    rPr33a = p33_after._element.get_or_add_rPr()
                    shd33a = OxmlElement('w:shd')
                    shd33a.set(qn('w:val'), 'clear')
                    shd33a.set(qn('w:color'), 'auto')
                    shd33a.set(qn('w:fill'), 'FFFED5')
                    rPr33a.append(shd33a)

                set_entry_box(paragraphs[34], "")  # blank entry box

                # Now strikethrough the NR NOT CONSULTED section (paras 24-36)
                for i in range(24, 37):
                    if i < len(paragraphs):
                        strikethrough_para(paragraphs[i])
            else:
                # NR NOT CONSULTED - apply highlighting to consulted section first, then strikethrough
                print("DEBUG: NR NOT CONSULTED section executing")
                print(f"DEBUG: nc_option_a={self.nc_option_a.isChecked()}, nc_option_b={self.nc_option_b.isChecked()}, nc_option_c={self.nc_option_c.isChecked()}")

                # Print paragraph contents to find correct indices
                for i in range(24, 40):
                    if i < len(paragraphs):
                        txt = ''.join(run.text for run in paragraphs[i].runs)[:50]
                        print(f"DEBUG: Para {i}: {txt}")

                # Apply cream highlighting to NR CONSULTED paragraphs first
                highlight_yellow(paragraphs[17])
                set_entry_box(paragraphs[18], "")
                highlight_yellow(paragraphs[19])
                highlight_yellow(paragraphs[20])
                set_entry_box(paragraphs[21], "")
                highlight_yellow(paragraphs[22])
                highlight_yellow(paragraphs[23])

                # Now strikethrough the NR CONSULTED section
                for i in range(15, 24):
                    if i < len(paragraphs):
                        strikethrough_para(paragraphs[i])

                # ALWAYS apply cream highlighting to ALL NC option paragraphs FIRST
                # Para 25: (a) I have been unable to ascertain...
                # Para 26: (b) To the best of my knowledge and belief...
                highlight_yellow(paragraphs[25])
                highlight_yellow(paragraphs[26])

                # Determine which NC option is selected, then strikethrough the non-selected ones
                if self.nc_option_a.isChecked():
                    # Option (a) selected - strikethrough (b) and (c) sections
                    strikethrough_para(paragraphs[26])
                    for i in range(27, 36):
                        if i < len(paragraphs):
                            strikethrough_para(paragraphs[i])

                elif self.nc_option_b.isChecked():
                    # Option (b) selected - strikethrough (a) and (c) sections
                    strikethrough_para(paragraphs[25])
                    for i in range(27, 36):
                        if i < len(paragraphs):
                            strikethrough_para(paragraphs[i])

                elif self.nc_option_c.isChecked():
                    # Option (c) selected - format the whole (c) section with gold brackets
                    # Strikethrough (a) and (b)
                    strikethrough_para(paragraphs[25])
                    strikethrough_para(paragraphs[26])

                    # Para 27: (c) I understand that... - cream highlight with gold bracket at start
                    p27 = paragraphs[27]
                    p27_text = ''.join(run.text for run in p27.runs)
                    for run in p27.runs:
                        run.text = ""
                    while len(p27.runs) > 1:
                        p27._element.remove(p27.runs[-1]._element)
                    # Add opening gold bracket
                    if p27.runs:
                        p27.runs[0].text = '['
                        p27.runs[0].font.bold = True
                        p27.runs[0].font.color.rgb = BRACKET_COLOR
                        rPr27o = p27.runs[0]._element.get_or_add_rPr()
                        shd27o = OxmlElement('w:shd')
                        shd27o.set(qn('w:val'), 'clear')
                        shd27o.set(qn('w:color'), 'auto')
                        shd27o.set(qn('w:fill'), 'FFFED5')
                        rPr27o.append(shd27o)
                    else:
                        bracket_open = p27.add_run('[')
                        bracket_open.font.bold = True
                        bracket_open.font.color.rgb = BRACKET_COLOR
                        rPr27o = bracket_open._element.get_or_add_rPr()
                        shd27o = OxmlElement('w:shd')
                        shd27o.set(qn('w:val'), 'clear')
                        shd27o.set(qn('w:color'), 'auto')
                        shd27o.set(qn('w:fill'), 'FFFED5')
                        rPr27o.append(shd27o)
                    # Add the text content with cream highlight
                    p27_content = p27.add_run(p27_text)
                    p27_content.font.name = 'Arial'
                    p27_content.font.size = Pt(12)
                    rPr27c = p27_content._element.get_or_add_rPr()
                    shd27c = OxmlElement('w:shd')
                    shd27c.set(qn('w:val'), 'clear')
                    shd27c.set(qn('w:color'), 'auto')
                    shd27c.set(qn('w:fill'), 'FFFED5')
                    rPr27c.append(shd27c)

                    # Para 28: NR entry box with gold brackets
                    nc_nr_text = self.nc_nr_name.text().strip()
                    if self.nc_nr_address.text().strip():
                        nc_nr_text += ", " + self.nc_nr_address.text().strip()
                    if nc_nr_text.strip():
                        set_entry_box(paragraphs[28], nc_nr_text)
                    else:
                        set_entry_box(paragraphs[28], "")

                    # Para 29: "is" - cream highlight (rebuild to ensure highlighting works)
                    p29 = paragraphs[29]
                    p29_text = ''.join(run.text for run in p29.runs)
                    if not p29_text.strip():
                        p29_text = ' is'
                    for run in p29.runs:
                        run.text = ""
                    while len(p29.runs) > 1:
                        p29._element.remove(p29.runs[-1]._element)
                    if p29.runs:
                        p29.runs[0].text = p29_text
                        p29.runs[0].font.name = 'Arial'
                        p29.runs[0].font.size = Pt(12)
                        rPr29 = p29.runs[0]._element.get_or_add_rPr()
                        shd29 = OxmlElement('w:shd')
                        shd29.set(qn('w:val'), 'clear')
                        shd29.set(qn('w:color'), 'auto')
                        shd29.set(qn('w:fill'), 'FFFED5')
                        rPr29.append(shd29)
                    else:
                        is_run = p29.add_run(p29_text)
                        is_run.font.name = 'Arial'
                        is_run.font.size = Pt(12)
                        rPr29 = is_run._element.get_or_add_rPr()
                        shd29 = OxmlElement('w:shd')
                        shd29.set(qn('w:val'), 'clear')
                        shd29.set(qn('w:color'), 'auto')
                        shd29.set(qn('w:fill'), 'FFFED5')
                        rPr29.append(shd29)

                    # Para 30: (i) this patient's nearest relative...
                    # Para 31: (ii) authorised to exercise the functions...
                    # BOTH get cream highlighting first, then one gets gold bracket, other gets struck through

                    # First, apply cream highlighting to BOTH paragraphs
                    highlight_yellow(paragraphs[30])
                    highlight_yellow(paragraphs[31])

                    if self.nc_c_is_nr.isChecked():
                        # (i) is selected - rebuild para 30 with gold bracket, strikethrough para 31
                        p30 = paragraphs[30]
                        p30_text = ''.join(run.text for run in p30.runs)
                        for run in p30.runs:
                            run.text = ""
                        while len(p30.runs) > 1:
                            p30._element.remove(p30.runs[-1]._element)
                        if p30.runs:
                            p30.runs[0].text = p30_text
                            p30.runs[0].font.name = 'Arial'
                            p30.runs[0].font.size = Pt(12)
                            rPr30 = p30.runs[0]._element.get_or_add_rPr()
                            shd30 = OxmlElement('w:shd')
                            shd30.set(qn('w:val'), 'clear')
                            shd30.set(qn('w:color'), 'auto')
                            shd30.set(qn('w:fill'), 'FFFED5')
                            rPr30.append(shd30)
                        # Add closing gold bracket to para 30
                        p30_close = p30.add_run(']')
                        p30_close.font.bold = True
                        p30_close.font.color.rgb = BRACKET_COLOR
                        rPr30c = p30_close._element.get_or_add_rPr()
                        shd30c = OxmlElement('w:shd')
                        shd30c.set(qn('w:val'), 'clear')
                        shd30c.set(qn('w:color'), 'auto')
                        shd30c.set(qn('w:fill'), 'FFFED5')
                        rPr30c.append(shd30c)
                        # Strikethrough para 31 (already has cream highlight from above)
                        strikethrough_para(paragraphs[31])
                    else:
                        # (ii) is selected - rebuild para 31 with gold bracket, strikethrough para 30
                        p31 = paragraphs[31]
                        p31_text = ''.join(run.text for run in p31.runs)
                        for run in p31.runs:
                            run.text = ""
                        while len(p31.runs) > 1:
                            p31._element.remove(p31.runs[-1]._element)
                        if p31.runs:
                            p31.runs[0].text = p31_text
                            p31.runs[0].font.name = 'Arial'
                            p31.runs[0].font.size = Pt(12)
                            rPr31 = p31.runs[0]._element.get_or_add_rPr()
                            shd31 = OxmlElement('w:shd')
                            shd31.set(qn('w:val'), 'clear')
                            shd31.set(qn('w:color'), 'auto')
                            shd31.set(qn('w:fill'), 'FFFED5')
                            rPr31.append(shd31)
                        # Add closing gold bracket to para 31
                        p31_close = p31.add_run(']')
                        p31_close.font.bold = True
                        p31_close.font.color.rgb = BRACKET_COLOR
                        rPr31c = p31_close._element.get_or_add_rPr()
                        shd31c = OxmlElement('w:shd')
                        shd31c.set(qn('w:val'), 'clear')
                        shd31c.set(qn('w:color'), 'auto')
                        shd31c.set(qn('w:fill'), 'FFFED5')
                        rPr31c.append(shd31c)
                        # Strikethrough para 30 (already has cream highlight from above)
                        strikethrough_para(paragraphs[30])

                    # Para 33: "not reasonably practicable/would involve unreasonable delay" with gold brackets
                    p33 = paragraphs[33]
                    p33_text = ''.join(run.text for run in p33.runs)
                    for run in p33.runs:
                        run.text = ""
                    while len(p33.runs) > 1:
                        p33._element.remove(p33.runs[-1]._element)
                    # Find the phrase to bracket
                    if "not reasonably practicable" in p33_text:
                        idx = p33_text.index("not reasonably practicable")
                        before = p33_text[:idx]
                        end_phrase = "unreasonable delay"
                        if end_phrase in p33_text:
                            end_idx = p33_text.index(end_phrase) + len(end_phrase)
                            bracketed = p33_text[idx:end_idx]
                            after = p33_text[end_idx:]
                        else:
                            bracketed = p33_text[idx:]
                            after = ""
                    else:
                        before = p33_text
                        bracketed = ""
                        after = ""

                    if p33.runs:
                        p33.runs[0].text = before
                        p33.runs[0].font.name = 'Arial'
                        p33.runs[0].font.size = Pt(12)
                        # Add cream highlight to "before" text
                        rPr33b = p33.runs[0]._element.get_or_add_rPr()
                        shd33b = OxmlElement('w:shd')
                        shd33b.set(qn('w:val'), 'clear')
                        shd33b.set(qn('w:color'), 'auto')
                        shd33b.set(qn('w:fill'), 'FFFED5')
                        rPr33b.append(shd33b)
                    if bracketed:
                        p33_bo = p33.add_run('[')
                        p33_bo.font.bold = True
                        p33_bo.font.color.rgb = BRACKET_COLOR
                        rPr33bo = p33_bo._element.get_or_add_rPr()
                        shd33bo = OxmlElement('w:shd')
                        shd33bo.set(qn('w:val'), 'clear')
                        shd33bo.set(qn('w:color'), 'auto')
                        shd33bo.set(qn('w:fill'), 'FFFED5')
                        rPr33bo.append(shd33bo)
                        p33_bc = p33.add_run(bracketed)
                        p33_bc.font.name = 'Arial'
                        p33_bc.font.size = Pt(12)
                        rPr33bc = p33_bc._element.get_or_add_rPr()
                        shd33bc = OxmlElement('w:shd')
                        shd33bc.set(qn('w:val'), 'clear')
                        shd33bc.set(qn('w:color'), 'auto')
                        shd33bc.set(qn('w:fill'), 'FFFED5')
                        rPr33bc.append(shd33bc)
                        p33_bcl = p33.add_run(']')
                        p33_bcl.font.bold = True
                        p33_bcl.font.color.rgb = BRACKET_COLOR
                        rPr33bcl = p33_bcl._element.get_or_add_rPr()
                        shd33bcl = OxmlElement('w:shd')
                        shd33bcl.set(qn('w:val'), 'clear')
                        shd33bcl.set(qn('w:color'), 'auto')
                        shd33bcl.set(qn('w:fill'), 'FFFED5')
                        rPr33bcl.append(shd33bcl)
                    if after:
                        p33_after = p33.add_run(after)
                        p33_after.font.name = 'Arial'
                        p33_after.font.size = Pt(12)
                        # Add cream highlight to "after" text
                        rPr33a = p33_after._element.get_or_add_rPr()
                        shd33a = OxmlElement('w:shd')
                        shd33a.set(qn('w:val'), 'clear')
                        shd33a.set(qn('w:color'), 'auto')
                        shd33a.set(qn('w:fill'), 'FFFED5')
                        rPr33a.append(shd33a)

                    # Para 34: Blank gold bracket placeholder OR reason text
                    if self.nc_reason.toPlainText():
                        set_entry_box(paragraphs[34], self.nc_reason.toPlainText())
                    else:
                        set_entry_box(paragraphs[34], "")

            # Debug: Print paragraphs 35-55 to find correct indices
            print("DEBUG: Paragraphs 35-55:")
            for i in range(35, min(56, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)[:60]
                print(f"DEBUG: Para {i}: {txt}")

            # Format para 37 "[If you need to continue...]" with checkbox - above "The remainder..."
            para37 = paragraphs[37]
            for run in para37.runs:
                run.text = ""
            while len(para37.runs) > 1:
                para37._element.remove(para37.runs[-1]._element)
            if para37.runs:
                para37.runs[0].text = '[If you need to continue on a separate sheet please indicate here '
                para37.runs[0].font.name = 'Arial'
                para37.runs[0].font.size = Pt(12)
            else:
                para37.add_run('[If you need to continue on a separate sheet please indicate here ')
            # Add gold checkbox [ ]
            cb37_open = para37.add_run('[')
            cb37_open.font.bold = True
            cb37_open.font.color.rgb = BRACKET_COLOR
            rPr37_o = cb37_open._element.get_or_add_rPr()
            shd37_o = OxmlElement('w:shd')
            shd37_o.set(qn('w:val'), 'clear')
            shd37_o.set(qn('w:color'), 'auto')
            shd37_o.set(qn('w:fill'), 'FFFED5')
            rPr37_o.append(shd37_o)
            cb37_space = para37.add_run(' ')
            rPr37_s = cb37_space._element.get_or_add_rPr()
            shd37_s = OxmlElement('w:shd')
            shd37_s.set(qn('w:val'), 'clear')
            shd37_s.set(qn('w:color'), 'auto')
            shd37_s.set(qn('w:fill'), 'FFFED5')
            rPr37_s.append(shd37_s)
            cb37_close = para37.add_run(']')
            cb37_close.font.bold = True
            cb37_close.font.color.rgb = BRACKET_COLOR
            rPr37_c = cb37_close._element.get_or_add_rPr()
            shd37_c = OxmlElement('w:shd')
            shd37_c.set(qn('w:val'), 'clear')
            shd37_c.set(qn('w:color'), 'auto')
            shd37_c.set(qn('w:fill'), 'FFFED5')
            rPr37_c.append(shd37_c)
            para37_end = para37.add_run(' and attach that sheet to this form]')
            para37_end.font.name = 'Arial'
            para37_end.font.size = Pt(12)

            # Date last seen - search for "I saw the patient on" to find correct paragraph
            date_para_idx = None
            for i in range(38, min(50, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "saw the patient" in txt.lower() or "[date]" in txt.lower():
                    date_para_idx = i + 1  # The entry box is the next paragraph
                    print(f"DEBUG: Found date paragraph at {i}, entry box at {i+1}")
                    break

            if date_para_idx is None:
                date_para_idx = 40  # fallback

            # Date last seen - gold bracketed entry box
            last_seen = self.last_seen_date.date().toString("dd MMMM yyyy")
            set_entry_box(paragraphs[date_para_idx], last_seen)

            # No acquaintance reason - search for it
            acq_para_idx = None
            for i in range(date_para_idx + 1, min(55, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "neither of the" in txt.lower() or "previous acquaintance" in txt.lower():
                    acq_para_idx = i + 1  # The entry box is the next paragraph
                    print(f"DEBUG: Found acquaintance paragraph at {i}, entry box at {i+1}")
                    break

            if acq_para_idx is None:
                acq_para_idx = 45  # fallback

            # No acquaintance reason - gold bracketed entry box
            if self.no_acquaintance_reason.toPlainText():
                set_entry_box(paragraphs[acq_para_idx], self.no_acquaintance_reason.toPlainText())
            else:
                set_entry_box(paragraphs[acq_para_idx], "")

            # Find signature line - search for "Signed"
            sig_para_idx = None
            for i in range(acq_para_idx + 1, min(60, len(paragraphs))):
                txt = ''.join(run.text for run in paragraphs[i].runs)
                if "signed" in txt.lower():
                    sig_para_idx = i
                    print(f"DEBUG: Found signature paragraph at {i}")
                    break

            if sig_para_idx is None:
                sig_para_idx = 49  # fallback

            # Format "[If you need to continue...]" just above signature (para 48)
            continue_para_idx = sig_para_idx - 1
            para_cont = paragraphs[continue_para_idx]
            for run in para_cont.runs:
                run.text = ""
            while len(para_cont.runs) > 1:
                para_cont._element.remove(para_cont.runs[-1]._element)
            if para_cont.runs:
                para_cont.runs[0].text = '[If you need to continue on a separate sheet please indicate here '
                para_cont.runs[0].font.name = 'Arial'
                para_cont.runs[0].font.size = Pt(12)
            else:
                para_cont.add_run('[If you need to continue on a separate sheet please indicate here ')
            # Add gold checkbox [ ]
            cb_open = para_cont.add_run('[')
            cb_open.font.bold = True
            cb_open.font.color.rgb = BRACKET_COLOR
            rPr_cbo = cb_open._element.get_or_add_rPr()
            shd_cbo = OxmlElement('w:shd')
            shd_cbo.set(qn('w:val'), 'clear')
            shd_cbo.set(qn('w:color'), 'auto')
            shd_cbo.set(qn('w:fill'), 'FFFED5')
            rPr_cbo.append(shd_cbo)
            cb_space = para_cont.add_run(' ')
            rPr_cbs = cb_space._element.get_or_add_rPr()
            shd_cbs = OxmlElement('w:shd')
            shd_cbs.set(qn('w:val'), 'clear')
            shd_cbs.set(qn('w:color'), 'auto')
            shd_cbs.set(qn('w:fill'), 'FFFED5')
            rPr_cbs.append(shd_cbs)
            cb_close = para_cont.add_run(']')
            cb_close.font.bold = True
            cb_close.font.color.rgb = BRACKET_COLOR
            rPr_cbc = cb_close._element.get_or_add_rPr()
            shd_cbc = OxmlElement('w:shd')
            shd_cbc.set(qn('w:val'), 'clear')
            shd_cbc.set(qn('w:color'), 'auto')
            shd_cbc.set(qn('w:fill'), 'FFFED5')
            rPr_cbc.append(shd_cbc)
            cont_end = para_cont.add_run(' and attach that sheet to this form]')
            cont_end.font.name = 'Arial'
            cont_end.font.size = Pt(12)

            # Signature line - "Signed[ ] Date[ ]" with gold brackets
            sig_date = self.signature_date.date().toString("dd MMMM yyyy")
            para_sig = paragraphs[sig_para_idx]
            for run in para_sig.runs:
                run.text = ""
            while len(para_sig.runs) > 1:
                para_sig._element.remove(para_sig.runs[-1]._element)
            if para_sig.runs:
                para_sig.runs[0].text = 'Signed'
                para_sig.runs[0].font.name = 'Arial'
                para_sig.runs[0].font.size = Pt(12)
            else:
                r = para_sig.add_run('Signed')
                r.font.name = 'Arial'
                r.font.size = Pt(12)
            sig_open = para_sig.add_run('[')
            sig_open.font.bold = True
            sig_open.font.color.rgb = BRACKET_COLOR
            rPr_so = sig_open._element.get_or_add_rPr()
            shd_so = OxmlElement('w:shd')
            shd_so.set(qn('w:val'), 'clear')
            shd_so.set(qn('w:color'), 'auto')
            shd_so.set(qn('w:fill'), 'FFFED5')
            rPr_so.append(shd_so)
            sig_content = para_sig.add_run('                                        ')
            rPr_sc = sig_content._element.get_or_add_rPr()
            shd_sc = OxmlElement('w:shd')
            shd_sc.set(qn('w:val'), 'clear')
            shd_sc.set(qn('w:color'), 'auto')
            shd_sc.set(qn('w:fill'), 'FFFED5')
            rPr_sc.append(shd_sc)
            sig_close = para_sig.add_run(']')
            sig_close.font.bold = True
            sig_close.font.color.rgb = BRACKET_COLOR
            rPr_scl = sig_close._element.get_or_add_rPr()
            shd_scl = OxmlElement('w:shd')
            shd_scl.set(qn('w:val'), 'clear')
            shd_scl.set(qn('w:color'), 'auto')
            shd_scl.set(qn('w:fill'), 'FFFED5')
            rPr_scl.append(shd_scl)
            date_label = para_sig.add_run('   Date')
            date_label.font.name = 'Arial'
            date_label.font.size = Pt(12)
            date_open = para_sig.add_run('[')
            date_open.font.bold = True
            date_open.font.color.rgb = BRACKET_COLOR
            rPr_do = date_open._element.get_or_add_rPr()
            shd_do = OxmlElement('w:shd')
            shd_do.set(qn('w:val'), 'clear')
            shd_do.set(qn('w:color'), 'auto')
            shd_do.set(qn('w:fill'), 'FFFED5')
            rPr_do.append(shd_do)
            date_content = para_sig.add_run(sig_date)
            rPr_dc = date_content._element.get_or_add_rPr()
            shd_dc = OxmlElement('w:shd')
            shd_dc.set(qn('w:val'), 'clear')
            shd_dc.set(qn('w:color'), 'auto')
            shd_dc.set(qn('w:fill'), 'FFFED5')
            rPr_dc.append(shd_dc)
            date_close = para_sig.add_run(']')
            date_close.font.bold = True
            date_close.font.color.rgb = BRACKET_COLOR
            rPr_dcl = date_close._element.get_or_add_rPr()
            shd_dcl = OxmlElement('w:shd')
            shd_dcl.set(qn('w:val'), 'clear')
            shd_dcl.set(qn('w:color'), 'auto')
            shd_dcl.set(qn('w:fill'), 'FFFED5')
            rPr_dcl.append(shd_dcl)

            doc.save(file_path)
            QMessageBox.information(self, "Export Complete", f"Form A6 exported to:\n{file_path}")

        except ImportError:
            QMessageBox.warning(self, "Export Error", "python-docx library required.")
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Export Error", f"Failed to export:\n{str(e)}\n\n{traceback.format_exc()}")

    def get_state(self) -> dict:
        return {
            "hospital_name": self.hospital_name.text(),
            "hospital_address": self.hospital_address.text(),
            "amhp_name": self.amhp_name.text(),
            "amhp_address": self.amhp_address.text(),
            "amhp_email": self.amhp_email.text(),
            "patient_name": self.patient_name.text(),
            "patient_address": self.patient_address.text(),
            "local_authority": self.local_authority.text(),
            "approved_by_same": self.approved_by_same.isChecked(),
            "approved_by_authority": self.approved_by_authority.text(),
            "nr_consulted": self.nr_consulted.isChecked(),
            "nr_option_a": self.nr_option_a.isChecked(),
            "nr_name": self.nr_name.text(),
            "nr_address": self.nr_address.text(),
            "nc_option_a": self.nc_option_a.isChecked(),
            "nc_option_b": self.nc_option_b.isChecked(),
            "nc_option_c": self.nc_option_c.isChecked(),
            "nc_nr_name": self.nc_nr_name.text(),
            "nc_nr_address": self.nc_nr_address.text(),
            "nc_c_is_nr": self.nc_c_is_nr.isChecked(),
            "nc_reason": self.nc_reason.toPlainText(),
            "last_seen_date": self.last_seen_date.date().toString("yyyy-MM-dd"),
            "no_acquaintance_reason": self.no_acquaintance_reason.toPlainText(),
            "signature_date": self.signature_date.date().toString("yyyy-MM-dd"),
        }

    def load_state(self, state: dict):
        if not state:
            return
        self.hospital_name.setText(state.get("hospital_name", ""))
        self.hospital_address.setText(state.get("hospital_address", ""))
        self.amhp_name.setText(state.get("amhp_name", ""))
        self.amhp_address.setText(state.get("amhp_address", ""))
        self.amhp_email.setText(state.get("amhp_email", ""))
        self.patient_name.setText(state.get("patient_name", ""))
        self.patient_address.setText(state.get("patient_address", ""))
        self.local_authority.setText(state.get("local_authority", ""))
        if state.get("approved_by_same", True):
            self.approved_by_same.setChecked(True)
        else:
            self.approved_by_different.setChecked(True)
        self.approved_by_authority.setText(state.get("approved_by_authority", ""))
        if state.get("nr_consulted", True):
            self.nr_consulted.setChecked(True)
        else:
            self.nr_not_consulted.setChecked(True)
        if state.get("nr_option_a", True):
            self.nr_option_a.setChecked(True)
        else:
            self.nr_option_b.setChecked(True)
        self.nr_name.setText(state.get("nr_name", ""))
        self.nr_address.setText(state.get("nr_address", ""))
        if state.get("nc_option_a", True):
            self.nc_option_a.setChecked(True)
        elif state.get("nc_option_b", False):
            self.nc_option_b.setChecked(True)
        elif state.get("nc_option_c", False):
            self.nc_option_c.setChecked(True)
        self.nc_nr_name.setText(state.get("nc_nr_name", ""))
        self.nc_nr_address.setText(state.get("nc_nr_address", ""))
        if state.get("nc_c_is_nr", True):
            self.nc_c_is_nr.setChecked(True)
        else:
            self.nc_c_is_authorised.setChecked(True)
        self.nc_reason.setPlainText(state.get("nc_reason", ""))
        if state.get("last_seen_date"):
            self.last_seen_date.setDate(QDate.fromString(state["last_seen_date"], "yyyy-MM-dd"))
        self.no_acquaintance_reason.setPlainText(state.get("no_acquaintance_reason", ""))
        if state.get("signature_date"):
            self.signature_date.setDate(QDate.fromString(state["signature_date"], "yyyy-MM-dd"))

    def _fill_patient_details(self, patient_info: dict):
        """Fill patient details from extracted demographics - only if fields are empty."""
        if not patient_info:
            return
        if patient_info.get("name") and not self.patient_name.text().strip():
            self.patient_name.setText(patient_info["name"])
            print(f"[A6Form] Set patient name: {patient_info['name']}")
